"""
Initial Setup: Create a senior thesis document with front matter and body chapters.
Task ID: writer_mt_066
Domain: libreoffice_writer

The document has:
- Front matter: Abstract, Acknowledgments, List of Abbreviations (Heading 1) with Roman numeral pages i-iv
- Body: 5 chapters (Heading 1) with Heading 2 sub-sections, Arabic pages 1-46
- NO Table of Contents (that is the agent's task)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_066'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_field(paragraph):
    """Add a PAGE field code to a paragraph for page numbering."""
    run1 = paragraph.add_run()
    fldChar1 = run1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run1._element.append(fldChar1)
    run2 = paragraph.add_run()
    instrText = run2._element.makeelement(qn('w:instrText'), {})
    instrText.text = ' PAGE '
    instrText.set(qn('xml:space'), 'preserve')
    run2._element.append(instrText)
    run3 = paragraph.add_run()
    fldChar3 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run3._element.append(fldChar3)


def set_page_number_format(section, fmt='decimal', start=None):
    """Set page number format on section. fmt: 'decimal', 'lowerRoman', 'upperRoman'."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = sectPr.makeelement(qn('w:pgNumType'), {})
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))


def add_body_text(doc, text, count=1):
    """Add realistic body paragraphs."""
    for _ in range(count):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'


def create_initial():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0

    # Configure heading styles
    for level in [1, 2]:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            hs.font.size = Pt(16)
            hs.font.bold = True
        else:
            hs.font.size = Pt(14)
            hs.font.bold = True

    # ============================================================
    # SECTION 1: FRONT MATTER (Roman numeral pages i-iv)
    # ============================================================
    section1 = doc.sections[0]
    section1.page_width = Inches(8.5)
    section1.page_height = Inches(11)
    section1.left_margin = Inches(1.5)
    section1.right_margin = Inches(1)
    section1.top_margin = Inches(1)
    section1.bottom_margin = Inches(1)
    set_page_number_format(section1, fmt='lowerRoman', start=1)

    # Footer with Roman page numbers for front matter
    footer1 = section1.footer
    footer1.is_linked_to_previous = False
    fp1 = footer1.paragraphs[0]
    fp1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number_field(fp1)

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    abstract_text = (
        "This thesis investigates the application of machine learning techniques "
        "to predictive maintenance in industrial manufacturing environments. "
        "As manufacturing processes become increasingly complex, traditional "
        "rule-based maintenance scheduling fails to account for the nuanced "
        "degradation patterns of modern equipment. Through analysis of sensor "
        "data collected from 147 CNC milling machines at three production "
        "facilities operated by Meridian Manufacturing Inc., this study develops "
        "and evaluates a hybrid deep learning framework combining convolutional "
        "neural networks with long short-term memory architectures."
    )
    add_body_text(doc, abstract_text)
    abstract_text2 = (
        "Results demonstrate that the proposed approach achieves a 94.3% "
        "prediction accuracy for bearing failures with a 72-hour lead time, "
        "representing a 23.7% improvement over conventional vibration analysis "
        "methods. The framework further reduces unplanned downtime by 31.2% "
        "and maintenance costs by $2.4 million annually across the three "
        "facilities studied. These findings contribute to the growing body of "
        "literature on Industry 4.0 applications and provide a practical "
        "methodology for implementing predictive maintenance systems in "
        "small to mid-sized manufacturing operations."
    )
    add_body_text(doc, abstract_text2)

    # Page break after Abstract
    doc.add_page_break()

    # --- Acknowledgments ---
    doc.add_heading('Acknowledgments', level=1)
    ack_text1 = (
        "I would like to express my deepest gratitude to my thesis advisor, "
        "Dr. Rebecca Morrison, for her invaluable guidance, patience, and "
        "encouragement throughout this research. Her expertise in industrial "
        "engineering and her commitment to rigorous scholarship have profoundly "
        "shaped this work and my development as a researcher."
    )
    add_body_text(doc, ack_text1)
    ack_text2 = (
        "I am also grateful to the members of my thesis committee, Dr. James "
        "Whitfield and Dr. Priya Chakraborty, for their insightful feedback "
        "and constructive criticism. Their perspectives from mechanical "
        "engineering and computer science respectively enriched the "
        "interdisciplinary nature of this study."
    )
    add_body_text(doc, ack_text2)
    ack_text3 = (
        "Special thanks go to the engineering team at Meridian Manufacturing "
        "Inc., particularly Chief Technology Officer Laura Engstrom and "
        "Senior Maintenance Engineer Carlos Reyes, who facilitated access to "
        "production data and provided essential domain expertise. This research "
        "would not have been possible without their collaboration."
    )
    add_body_text(doc, ack_text3)

    # Page break
    doc.add_page_break()

    # --- List of Abbreviations ---
    doc.add_heading('List of Abbreviations', level=1)
    abbreviations = [
        ("AI", "Artificial Intelligence"),
        ("ANN", "Artificial Neural Network"),
        ("CBM", "Condition-Based Maintenance"),
        ("CNC", "Computer Numerical Control"),
        ("CNN", "Convolutional Neural Network"),
        ("DL", "Deep Learning"),
        ("FFT", "Fast Fourier Transform"),
        ("GAN", "Generative Adversarial Network"),
        ("IoT", "Internet of Things"),
        ("LSTM", "Long Short-Term Memory"),
        ("MAE", "Mean Absolute Error"),
        ("ML", "Machine Learning"),
        ("MTBF", "Mean Time Between Failures"),
        ("MTTF", "Mean Time To Failure"),
        ("OEE", "Overall Equipment Effectiveness"),
        ("PCA", "Principal Component Analysis"),
        ("PdM", "Predictive Maintenance"),
        ("PM", "Preventive Maintenance"),
        ("RMSE", "Root Mean Square Error"),
        ("ROC", "Receiver Operating Characteristic"),
        ("RUL", "Remaining Useful Life"),
        ("SVM", "Support Vector Machine"),
    ]
    for abbr, full in abbreviations:
        p = doc.add_paragraph()
        run_abbr = p.add_run(f"{abbr}")
        run_abbr.bold = True
        run_abbr.font.size = Pt(12)
        run_abbr.font.name = 'Times New Roman'
        run_rest = p.add_run(f"\t{full}")
        run_rest.font.size = Pt(12)
        run_rest.font.name = 'Times New Roman'

    # ============================================================
    # SECTION 2: BODY (Arabic numeral pages 1-46)
    # ============================================================
    # Add a section break to start body on new page with Arabic numbering
    new_section_para = doc.add_paragraph()
    new_section_para.paragraph_format.page_break_before = False
    run = new_section_para.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)

    # We need to add a section break via XML
    last_para = doc.paragraphs[-1]
    sectPr = last_para._element.makeelement(qn('w:sectPr'), {})
    # Copy section properties from section1
    sectPr.append(sectPr.makeelement(qn('w:pgSz'), {
        qn('w:w'): str(section1.page_width),
        qn('w:h'): str(section1.page_height),
    }))
    sectPr.append(sectPr.makeelement(qn('w:pgMar'), {
        qn('w:left'): str(section1.left_margin),
        qn('w:right'): str(section1.right_margin),
        qn('w:top'): str(section1.top_margin),
        qn('w:bottom'): str(section1.bottom_margin),
        qn('w:header'): '720',
        qn('w:footer'): '720',
    }))
    sectPr.append(sectPr.makeelement(qn('w:pgNumType'), {
        qn('w:fmt'): 'lowerRoman',
        qn('w:start'): '1',
    }))
    sectPr.append(sectPr.makeelement(qn('w:type'), {
        qn('w:val'): 'nextPage',
    }))
    last_para._element.addnext(sectPr)

    # Actually, use python-docx section approach
    # Remove the manually added sectPr and use proper approach
    last_para._element.getparent().remove(sectPr)

    # Instead, add section break properly
    # We'll create a new section by adding sectPr to the last paragraph's pPr
    pPr = last_para._element.get_or_add_pPr()
    sect_pr_in_para = pPr.makeelement(qn('w:sectPr'), {})
    # page size
    pgSz = sect_pr_in_para.makeelement(qn('w:pgSz'), {
        qn('w:w'): str(int(Inches(8.5))),
        qn('w:h'): str(int(Inches(11))),
    })
    sect_pr_in_para.append(pgSz)
    # page margins
    pgMar = sect_pr_in_para.makeelement(qn('w:pgMar'), {
        qn('w:left'): str(int(Inches(1.5))),
        qn('w:right'): str(int(Inches(1))),
        qn('w:top'): str(int(Inches(1))),
        qn('w:bottom'): str(int(Inches(1))),
        qn('w:header'): '720',
        qn('w:footer'): '720',
    })
    sect_pr_in_para.append(pgMar)
    # Roman page numbering for this section
    pgNum = sect_pr_in_para.makeelement(qn('w:pgNumType'), {
        qn('w:fmt'): 'lowerRoman',
        qn('w:start'): '1',
    })
    sect_pr_in_para.append(pgNum)
    # Footer reference for this section
    pPr.append(sect_pr_in_para)

    # Now the body section (section 2) starts
    # After saving, doc.sections[1] will be the body section
    # We'll configure it after adding content

    # --- Chapter 1: Introduction ---
    doc.add_heading('Chapter 1: Introduction', level=1)

    doc.add_heading('1.1 Background and Motivation', level=2)
    add_body_text(doc, (
        "The manufacturing sector has undergone significant transformation over "
        "the past two decades, driven by advances in sensor technology, data "
        "analytics, and automation. The concept of Industry 4.0, first introduced "
        "at the Hannover Messe in 2011, envisions smart factories where "
        "cyber-physical systems monitor production processes in real time, "
        "enabling unprecedented levels of efficiency and quality control."
    ))
    add_body_text(doc, (
        "Equipment maintenance represents one of the most critical operational "
        "challenges in modern manufacturing. According to a 2023 report by "
        "Deloitte, unplanned downtime costs industrial manufacturers an estimated "
        "$50 billion annually in the United States alone. Traditional maintenance "
        "strategies fall into two broad categories: reactive maintenance, where "
        "equipment is repaired after failure, and preventive maintenance, where "
        "servicing occurs at predetermined intervals regardless of actual "
        "equipment condition."
    ))
    add_body_text(doc, (
        "Predictive maintenance (PdM) offers a third approach that leverages "
        "real-time monitoring data to forecast equipment failures before they "
        "occur. By analyzing patterns in vibration, temperature, acoustic "
        "emission, and other sensor signals, PdM systems can identify early "
        "indicators of degradation and schedule maintenance optimally. This "
        "approach minimizes both unnecessary preventive interventions and "
        "costly reactive repairs."
    ))

    doc.add_heading('1.2 Problem Statement', level=2)
    add_body_text(doc, (
        "Despite the theoretical advantages of predictive maintenance, practical "
        "implementation remains challenging for several reasons. First, "
        "manufacturing environments generate enormous volumes of sensor data "
        "that must be processed and analyzed in near real-time. Second, "
        "degradation patterns vary significantly across equipment types, "
        "operating conditions, and production loads. Third, labeled failure data "
        "is inherently scarce because well-maintained equipment fails "
        "infrequently, creating severe class imbalance in training datasets."
    ))
    add_body_text(doc, (
        "Current approaches to machine learning-based predictive maintenance "
        "often rely on handcrafted features extracted from raw sensor signals. "
        "While effective in controlled laboratory settings, these feature "
        "engineering pipelines frequently fail to generalize across different "
        "equipment configurations and operating regimes. This thesis addresses "
        "the need for an automated, end-to-end deep learning framework that "
        "can learn discriminative features directly from raw sensor data."
    ))

    doc.add_heading('1.3 Research Objectives', level=2)
    add_body_text(doc, (
        "The primary objective of this research is to develop and evaluate a "
        "hybrid deep learning framework for predictive maintenance in CNC "
        "milling operations. Specifically, this thesis aims to: (1) design a "
        "CNN-LSTM architecture that extracts spatial and temporal features from "
        "multi-channel sensor data; (2) develop a transfer learning strategy "
        "that enables model adaptation across different machine types; (3) "
        "evaluate the framework using real-world production data from Meridian "
        "Manufacturing Inc.; and (4) provide practical guidelines for deploying "
        "PdM systems in small to mid-sized manufacturing operations."
    ))

    doc.add_heading('1.4 Thesis Organization', level=2)
    add_body_text(doc, (
        "The remainder of this thesis is organized as follows. Chapter 2 "
        "provides a comprehensive review of related literature in predictive "
        "maintenance, deep learning, and industrial IoT. Chapter 3 describes "
        "the proposed methodology, including the CNN-LSTM architecture, data "
        "preprocessing pipeline, and training procedures. Chapter 4 presents "
        "the experimental setup and results. Chapter 5 discusses the findings, "
        "limitations, and directions for future research."
    ))

    # --- Chapter 2: Literature Review ---
    doc.add_heading('Chapter 2: Literature Review', level=1)

    doc.add_heading('2.1 Traditional Maintenance Strategies', level=2)
    add_body_text(doc, (
        "Maintenance management has evolved considerably since the early days of "
        "industrialization. Moubray (1997) identified three generations of "
        "maintenance philosophy: the first generation (pre-1950s) relied "
        "entirely on fix-it-when-it-breaks approaches; the second generation "
        "(1950s-1980s) introduced scheduled preventive maintenance based on "
        "statistical failure models; and the third generation (1980s-present) "
        "emphasizes condition monitoring and reliability-centered maintenance."
    ))
    add_body_text(doc, (
        "Preventive maintenance (PM) scheduling typically follows either "
        "age-based or block-based replacement policies. In age-based policies, "
        "components are replaced after accumulating a predetermined number of "
        "operating hours. Block-based policies replace components at fixed "
        "calendar intervals regardless of usage. While PM reduces the frequency "
        "of catastrophic failures, it often results in premature replacement of "
        "components that still have significant remaining useful life."
    ))
    add_body_text(doc, (
        "Condition-based maintenance (CBM) emerged as a more sophisticated "
        "approach that monitors actual equipment condition through periodic or "
        "continuous inspection. ISO 17359:2018 provides guidelines for CBM "
        "implementation, recommending a systematic process of data acquisition, "
        "data processing, health assessment, and maintenance decision-making. "
        "CBM has been widely adopted in industries such as aviation, power "
        "generation, and oil refining, where equipment failures carry high "
        "safety and financial consequences."
    ))

    doc.add_heading('2.2 Machine Learning for Predictive Maintenance', level=2)
    add_body_text(doc, (
        "The application of machine learning to maintenance prediction has grown "
        "rapidly since the mid-2000s. Early approaches used classical algorithms "
        "such as support vector machines (Widodo and Yang, 2007), random forests "
        "(Loutas et al., 2011), and hidden Markov models (Tobon-Mejia et al., "
        "2012) to classify equipment health states from extracted features. "
        "These methods demonstrated promising results on benchmark datasets but "
        "required extensive domain expertise for feature engineering."
    ))
    add_body_text(doc, (
        "More recent work has explored deep learning architectures that can "
        "automatically learn hierarchical feature representations from raw "
        "sensor data. Li et al. (2018) applied convolutional neural networks to "
        "bearing fault diagnosis, achieving 99.2% accuracy on the Case Western "
        "Reserve University bearing dataset. Zhang et al. (2019) proposed a "
        "deep transfer learning approach for cross-domain fault diagnosis, "
        "demonstrating that features learned from one type of rotating machinery "
        "could be transferred to another with limited fine-tuning data."
    ))

    doc.add_heading('2.3 Deep Learning Architectures', level=2)
    add_body_text(doc, (
        "Convolutional neural networks (CNNs) have proven particularly effective "
        "for spatial feature extraction from time-frequency representations of "
        "vibration signals. Short-time Fourier transforms, wavelet transforms, "
        "and Hilbert-Huang transforms convert one-dimensional time series into "
        "two-dimensional spectrograms that CNN architectures can process "
        "similarly to images. This approach captures both frequency content and "
        "temporal evolution of fault signatures."
    ))
    add_body_text(doc, (
        "Long short-term memory (LSTM) networks excel at modeling temporal "
        "dependencies in sequential data. Unlike standard recurrent neural "
        "networks, LSTMs use gating mechanisms to selectively retain or discard "
        "information over long time horizons, making them well-suited for "
        "capturing gradual degradation trends. Zhao et al. (2017) demonstrated "
        "that LSTM networks outperform traditional regression methods for "
        "remaining useful life estimation on the NASA C-MAPSS turbofan dataset."
    ))

    doc.add_heading('2.4 Industrial IoT and Edge Computing', level=2)
    add_body_text(doc, (
        "The proliferation of low-cost sensors and wireless communication "
        "protocols has enabled comprehensive monitoring of manufacturing "
        "equipment at unprecedented scale. Modern CNC machines can generate "
        "thousands of data points per second from accelerometers, current "
        "sensors, acoustic emission transducers, and thermal imaging cameras. "
        "However, transmitting all raw data to cloud servers for processing "
        "introduces latency, bandwidth, and privacy concerns."
    ))
    add_body_text(doc, (
        "Edge computing architectures address these challenges by performing "
        "initial data processing and feature extraction at or near the data "
        "source. Lightweight models deployed on edge devices can provide "
        "real-time anomaly detection while transmitting only relevant features "
        "or alerts to centralized servers for more sophisticated analysis. "
        "This hybrid architecture balances the need for low-latency local "
        "decision-making with the computational resources available in the cloud."
    ))

    # --- Chapter 3: Methodology ---
    doc.add_heading('Chapter 3: Methodology', level=1)

    doc.add_heading('3.1 Data Collection', level=2)
    add_body_text(doc, (
        "Data was collected from 147 CNC milling machines operating across "
        "three Meridian Manufacturing facilities in Detroit, Michigan; "
        "Guadalajara, Mexico; and Stuttgart, Germany. Each machine was "
        "instrumented with a suite of sensors including tri-axial "
        "accelerometers (PCB Piezotronics 356A17) mounted on the spindle "
        "housing, current transformers on the spindle motor drive, acoustic "
        "emission sensors (Physical Acoustics R15I-AST), and K-type "
        "thermocouples at critical bearing locations."
    ))
    add_body_text(doc, (
        "Sensor data was sampled at 20 kHz for vibration and acoustic emission "
        "channels and 1 kHz for current and temperature channels. A National "
        "Instruments CompactDAQ system at each machine aggregated multi-channel "
        "data and transmitted 10-second windows to a local edge server every "
        "30 seconds. Over the 18-month data collection period from January 2024 "
        "to June 2025, the system recorded approximately 2.3 terabytes of raw "
        "sensor data, including 284 documented bearing failure events with "
        "detailed maintenance logs."
    ))

    doc.add_heading('3.2 Data Preprocessing', level=2)
    add_body_text(doc, (
        "Raw sensor signals were preprocessed through a multi-stage pipeline "
        "designed to remove noise, normalize signal amplitudes, and create "
        "input representations suitable for the deep learning framework. First, "
        "signals were bandpass filtered using a fourth-order Butterworth filter "
        "with cutoff frequencies of 10 Hz and 8 kHz to remove DC offset and "
        "high-frequency noise above the Nyquist frequency of interest."
    ))
    add_body_text(doc, (
        "Time-frequency representations were generated using continuous wavelet "
        "transforms with Morlet wavelets. Each 10-second data window was "
        "transformed into a 128 x 256 scalogram capturing frequency content "
        "from 10 Hz to 8 kHz with logarithmic frequency spacing. Multi-channel "
        "scalograms from the three accelerometer axes and the acoustic emission "
        "sensor were stacked to create four-channel input tensors analogous to "
        "multi-spectral images."
    ))

    doc.add_heading('3.3 CNN-LSTM Architecture', level=2)
    add_body_text(doc, (
        "The proposed hybrid architecture consists of three main components: "
        "a convolutional feature extractor, a temporal sequence encoder, and a "
        "classification head. The convolutional component uses a modified "
        "ResNet-18 backbone with four residual blocks, each containing two "
        "3x3 convolutional layers with batch normalization and ReLU activation. "
        "The initial 7x7 convolution is replaced with three successive 3x3 "
        "convolutions to better capture fine-grained spectral features."
    ))
    add_body_text(doc, (
        "Feature maps from the final convolutional layer are pooled using "
        "global average pooling to produce a 512-dimensional feature vector "
        "for each input window. These vectors are arranged chronologically and "
        "fed into a two-layer bidirectional LSTM with 256 hidden units per "
        "direction. The LSTM processes sequences of 24 consecutive windows "
        "(representing 4 minutes of operation) to capture temporal degradation "
        "trends that span multiple measurement intervals."
    ))

    doc.add_heading('3.4 Training Procedure', level=2)
    add_body_text(doc, (
        "The model was trained using a two-phase approach. In the first phase, "
        "the CNN feature extractor was pre-trained on a self-supervised "
        "contrastive learning task using all available unlabeled data. Pairs of "
        "augmented scalograms from the same time window were treated as positive "
        "examples, while scalograms from different machines at different times "
        "served as negative examples. This pre-training phase used the NT-Xent "
        "loss function with a temperature parameter of 0.07 and trained for "
        "200 epochs with a batch size of 512."
    ))
    add_body_text(doc, (
        "In the second phase, the full CNN-LSTM pipeline was fine-tuned on "
        "labeled failure data using focal loss (Lin et al., 2017) to address "
        "class imbalance. The healthy-to-faulty ratio in the training data was "
        "approximately 50:1, and focal loss with gamma=2.0 effectively down-"
        "weighted the contribution of well-classified healthy examples. Training "
        "used the Adam optimizer with an initial learning rate of 1e-4, cosine "
        "annealing schedule, and early stopping based on validation AUROC with "
        "a patience of 15 epochs."
    ))

    doc.add_heading('3.5 Transfer Learning Strategy', level=2)
    add_body_text(doc, (
        "To evaluate the generalizability of the learned representations, a "
        "transfer learning protocol was designed with three scenarios: (1) "
        "within-facility transfer, where models trained on one set of machines "
        "are applied to different machines at the same facility; (2) cross-"
        "facility transfer, where models are transferred between facilities "
        "with different environmental conditions; and (3) cross-machine-type "
        "transfer, where models trained on one CNC machine model are applied to "
        "a different model with different spindle configurations."
    ))

    # --- Chapter 4: Results ---
    doc.add_heading('Chapter 4: Results', level=1)

    doc.add_heading('4.1 Baseline Comparisons', level=2)
    add_body_text(doc, (
        "The proposed CNN-LSTM framework was compared against five baseline "
        "methods: (1) traditional vibration analysis using ISO 10816 velocity "
        "thresholds; (2) support vector machine (SVM) with handcrafted "
        "time-domain and frequency-domain features; (3) random forest with "
        "the same feature set; (4) standalone CNN without temporal modeling; "
        "and (5) standalone LSTM with raw signal input. All methods were "
        "evaluated using five-fold cross-validation with stratified splitting "
        "to maintain class balance across folds."
    ))
    add_body_text(doc, (
        "Table 4.1 summarizes the classification performance of all methods. "
        "The CNN-LSTM framework achieved the highest overall accuracy of 94.3% "
        "with a 72-hour prediction horizon, compared to 70.6% for ISO "
        "thresholds, 81.2% for SVM, 83.7% for random forest, 89.1% for "
        "standalone CNN, and 86.4% for standalone LSTM. The hybrid model also "
        "achieved the best F1 score (0.891) and AUROC (0.967), demonstrating "
        "superior discriminative ability across different operating conditions."
    ))

    doc.add_heading('4.2 Remaining Useful Life Estimation', level=2)
    add_body_text(doc, (
        "Beyond binary classification, the framework was evaluated on remaining "
        "useful life (RUL) estimation, which provides more actionable "
        "information for maintenance planning. The LSTM decoder was modified "
        "to output a continuous RUL estimate in hours rather than a binary "
        "classification. Using mean absolute error (MAE) as the evaluation "
        "metric, the CNN-LSTM achieved an MAE of 18.4 hours compared to 42.7 "
        "hours for the SVM regression baseline and 27.1 hours for standalone "
        "LSTM."
    ))
    add_body_text(doc, (
        "Figure 4.3 shows RUL prediction trajectories for five representative "
        "bearing failure cases. In most cases, the model correctly identifies "
        "the onset of degradation approximately 5-7 days before actual failure "
        "and provides increasingly accurate RUL estimates as the failure "
        "approaches. The model tends to overestimate RUL in the early stages "
        "of degradation and slightly underestimate in the final 24 hours, "
        "which is a conservative behavior preferred for maintenance scheduling."
    ))

    doc.add_heading('4.3 Transfer Learning Results', level=2)
    add_body_text(doc, (
        "Within-facility transfer yielded the strongest results, with accuracy "
        "dropping by only 2.1 percentage points on average when applying models "
        "to unseen machines at the same facility. Cross-facility transfer "
        "showed a larger performance gap of 8.7 percentage points, primarily "
        "attributable to differences in ambient temperature and humidity between "
        "the Detroit and Guadalajara facilities. However, fine-tuning with just "
        "50 labeled examples from the target facility recovered 73% of the "
        "performance gap."
    ))

    doc.add_heading('4.4 Operational Impact Analysis', level=2)
    add_body_text(doc, (
        "Deploying the CNN-LSTM framework at Meridian's Detroit facility over a "
        "six-month pilot period resulted in measurable operational improvements. "
        "Unplanned downtime decreased from an average of 47.3 hours per month "
        "to 32.5 hours per month, a 31.2% reduction. Maintenance labor costs "
        "decreased by 18.7% as technicians could plan interventions during "
        "scheduled production breaks rather than responding to emergency calls. "
        "The estimated annual cost savings for the single facility totaled "
        "$823,000, with projected savings of $2.4 million across all three "
        "facilities."
    ))

    # --- Chapter 5: Discussion and Conclusion ---
    doc.add_heading('Chapter 5: Discussion and Conclusion', level=1)

    doc.add_heading('5.1 Summary of Contributions', level=2)
    add_body_text(doc, (
        "This thesis has made several contributions to the field of predictive "
        "maintenance. First, it has demonstrated that hybrid CNN-LSTM "
        "architectures can effectively learn discriminative features from raw "
        "multi-channel sensor data without manual feature engineering. Second, "
        "it has shown that self-supervised pre-training on unlabeled data "
        "significantly improves classification performance in data-scarce "
        "industrial settings. Third, it has provided empirical evidence that "
        "transfer learning can reduce the data requirements for deploying PdM "
        "systems on new equipment."
    ))

    doc.add_heading('5.2 Limitations', level=2)
    add_body_text(doc, (
        "Several limitations of this work should be acknowledged. The study "
        "focused exclusively on bearing failures in CNC milling machines, and "
        "the generalizability of the framework to other failure modes such as "
        "tool wear, spindle misalignment, or coolant system degradation requires "
        "further investigation. Additionally, the 72-hour prediction horizon, "
        "while practically useful, may not be sufficient for facilities that "
        "require longer lead times for parts procurement."
    ))

    doc.add_heading('5.3 Future Work', level=2)
    add_body_text(doc, (
        "Several promising directions for future research emerge from this work. "
        "First, the integration of physics-informed neural networks could "
        "incorporate domain knowledge about bearing degradation mechanics into "
        "the learning framework, potentially improving interpretability and "
        "extrapolation to unseen operating conditions. Second, federated "
        "learning approaches could enable collaborative model training across "
        "multiple manufacturing organizations without sharing proprietary "
        "production data."
    ))
    add_body_text(doc, (
        "Third, the framework could be extended to multi-task learning settings "
        "where a single model simultaneously predicts failures across multiple "
        "components and failure modes. Finally, the integration of natural "
        "language processing for automatic extraction of maintenance insights "
        "from unstructured maintenance logs could provide additional training "
        "signals and enable more comprehensive health assessment of complex "
        "manufacturing systems."
    ))

    doc.add_heading('5.4 Concluding Remarks', level=2)
    add_body_text(doc, (
        "As manufacturing continues to evolve toward greater automation and "
        "digitalization, predictive maintenance will play an increasingly "
        "central role in operational excellence. The framework presented in "
        "this thesis demonstrates that deep learning approaches can deliver "
        "significant improvements over traditional methods while remaining "
        "practical for deployment in real-world manufacturing environments. "
        "The combination of automated feature learning, transfer capabilities, "
        "and demonstrated cost savings makes a compelling case for broader "
        "adoption of AI-driven maintenance strategies in the manufacturing "
        "sector."
    ))

    # ============================================================
    # Configure body section (section 2) page numbering
    # ============================================================
    # After all content is added, configure the final section (body section)
    if len(doc.sections) > 1:
        section2 = doc.sections[-1]
    else:
        # Only one section exists - the sectPr in para approach creates the split
        section2 = doc.sections[0]

    # Set Arabic page numbering starting at 1 for body section
    set_page_number_format(section2, fmt='decimal', start=1)
    section2.page_width = Inches(8.5)
    section2.page_height = Inches(11)
    section2.left_margin = Inches(1.5)
    section2.right_margin = Inches(1)
    section2.top_margin = Inches(1)
    section2.bottom_margin = Inches(1)

    # Footer with Arabic page numbers for body
    footer2 = section2.footer
    footer2.is_linked_to_previous = False
    fp2 = footer2.paragraphs[0]
    fp2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number_field(fp2)

    # Save document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
