"""
Initial Setup: Customer service script with 7 sentences in a continuous first paragraph
Task ID: osworld_writer_spacing_008
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_spacing_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- First paragraph: 7 scripted response sentences in a continuous block ---
    # These are the sentences the agent must split apart with empty paragraphs
    first_para = doc.add_paragraph(
        "Thank you for contacting us today. "
        "I understand your concern completely. "
        "Let me pull up your account right now. "
        "I can see the issue you are referring to and I sincerely apologize for the inconvenience. "
        "I will make sure this gets resolved for you as quickly as possible. "
        "Is there anything else I can help clarify while I work on this? "
        "Thank you for your patience and for being a valued customer."
    )

    # --- Second paragraph: Escalation procedure 1 (leave untouched) ---
    escalation_1 = doc.add_paragraph(
        "Escalation Procedure Level 1: If the customer remains unsatisfied after standard resolution attempts, "
        "politely inform them that you are escalating the case to a senior representative. "
        "Provide the customer with a case reference number and an estimated callback window of 24-48 business hours. "
        "Document all details in the CRM system before transferring."
    )

    # --- Third paragraph: Escalation procedure 2 (leave untouched) ---
    escalation_2 = doc.add_paragraph(
        "Escalation Procedure Level 2: For unresolved billing disputes or service outages affecting multiple accounts, "
        "immediately notify the on-call supervisor and open a priority ticket in the incident management system. "
        "Send the customer a written confirmation email within 30 minutes outlining the steps being taken. "
        "Follow up every 4 hours until the issue is fully resolved and the customer confirms satisfaction."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
