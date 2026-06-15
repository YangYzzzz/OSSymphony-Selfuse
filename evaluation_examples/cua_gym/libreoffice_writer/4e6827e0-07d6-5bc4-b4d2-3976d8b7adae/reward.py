"""
Reward Script: Apply character backgrounds and bold/colored dollar amounts
Task ID: writer_txtfmt_059
Domain: libreoffice_writer
Scoring:
  Component 1: Para 1 has light blue (#ADD8E6) character background for all runs (0.30)
  Component 2: Para 2 has light green (#90EE90) character background for all runs (0.30)
  Component 3: Dollar amounts ($45.2M, $31.8M, $2.4M) are bold AND dark green (#006400) (0.40)
Total: 1.0
"""

import os

# python-docx for .docx reading
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_059'

# File is on the Desktop
FILE_PATH = f'{WORKDIR}/Desktop/exec_summary.docx'

# Target colors (lowercase hex, no #)
PARA1_BG = 'add8e6'   # light blue (#ADD8E6)
PARA2_BG = '90ee90'   # light green (#90EE90)
DOLLAR_COLOR = '006400'  # dark green (#006400)

# Dollar amounts that must be bold and dark green
DOLLAR_AMOUNTS = ['$45.2 million', '$31.8 million', '$2.4 million']


def get_run_shd_fill(run):
    """
    Extracts the w:shd w:fill attribute from a run's rPr.
    Returns the fill value (lowercase hex string) or None if not set.
    """
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        return None
    shd = rPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    if fill:
        return fill.lower()
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Validate basic structure: should have at least 3 paragraphs
    if len(doc.paragraphs) < 3:
        print(f"CRITICAL: Expected at least 3 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Para index 0: Title ("Executive Summary")
    # Para index 1: First executive summary paragraph (Acme Corporation...)
    # Para index 2: Second executive summary paragraph (Operating expenses...)
    para1 = doc.paragraphs[1]
    para2 = doc.paragraphs[2]

    # Component 1: Para 1 has light blue (#ADD8E6) character background
    # on ALL runs (0.30 points)
    # This FAILS on initial (no shd element) → PASSES on golden (shd fill=ADD8E6)
    try:
        para1_runs = para1.runs
        if len(para1_runs) == 0:
            print("FAIL: Component 1 — Para 1 has no runs")
        else:
            all_blue = True
            for i, run in enumerate(para1_runs):
                fill = get_run_shd_fill(run)
                if fill != PARA1_BG:
                    all_blue = False
                    print(f"FAIL: Component 1 — Run[{i}] shd fill={repr(fill)}, expected '{PARA1_BG}'")
            if all_blue:
                print(f"PASS: Component 1 — All {len(para1_runs)} runs in Para 1 have light blue background (#ADD8E6) (0.30 pts)")
                total_score += 0.30
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Para 2 has light green (#90EE90) character background
    # on ALL runs (0.30 points)
    # This FAILS on initial (no shd element) → PASSES on golden (shd fill=90EE90)
    try:
        para2_runs = para2.runs
        if len(para2_runs) == 0:
            print("FAIL: Component 2 — Para 2 has no runs")
        else:
            all_green = True
            for i, run in enumerate(para2_runs):
                fill = get_run_shd_fill(run)
                if fill != PARA2_BG:
                    all_green = False
                    print(f"FAIL: Component 2 — Run[{i}] shd fill={repr(fill)}, expected '{PARA2_BG}'")
            if all_green:
                print(f"PASS: Component 2 — All {len(para2_runs)} runs in Para 2 have light green background (#90EE90) (0.30 pts)")
                total_score += 0.30
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Dollar amounts are bold AND dark green (#006400) (0.40 points)
    # Check each of the 3 dollar amounts: $45.2 million, $31.8 million, $2.4 million
    # This FAILS on initial (single run per para, no bold/color) → PASSES on golden
    try:
        dollar_results = {}
        for amount in DOLLAR_AMOUNTS:
            dollar_results[amount] = False

        # Search all runs in para1 and para2 for dollar amount text
        for para in [para1, para2]:
            for run in para.runs:
                for amount in DOLLAR_AMOUNTS:
                    if run.text.strip() == amount:
                        # Check bold: True means explicitly bold, None means inherit
                        is_bold = run.bold is True
                        # Check color: must be dark green #006400
                        color_rgb = run.font.color.rgb
                        is_dark_green = (color_rgb is not None and str(color_rgb).lower() == DOLLAR_COLOR)
                        if is_bold and is_dark_green:
                            dollar_results[amount] = True
                            print(f"PASS: Component 3 — '{amount}' is bold and dark green (#006400)")
                        else:
                            print(f"FAIL: Component 3 — '{amount}': bold={is_bold}, color={color_rgb} (expected bold=True, color=#006400)")

        # Partial credit: each dollar amount is worth 0.40/3 ≈ 0.133
        # All 3 must pass for full points; award all-or-nothing here for simplicity
        passed_count = sum(1 for v in dollar_results.values() if v)
        if passed_count == 3:
            print(f"PASS: Component 3 — All 3 dollar amounts are bold and dark green (0.40 pts)")
            total_score += 0.40
        elif passed_count > 0:
            partial = round(passed_count * (0.40 / 3), 2)
            print(f"PARTIAL: Component 3 — {passed_count}/3 dollar amounts correct ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No dollar amounts have bold + dark green formatting (0.0 pts)")

    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint: test against canonical artifact path on VM
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
