"""
Initial Setup: Apply text formatting to executive summary document
Task ID: writer_txtfmt_059
Domain: libreoffice_writer

Creates an executive summary document with plain text (no character backgrounds,
no bold/colored dollar amounts) as the pre-task state.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_059'
OUTPUT = f'{WORKDIR}/Desktop/exec_summary.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    desktop_dir = f'{WORKDIR}/Desktop'
    os.makedirs(desktop_dir, exist_ok=True)

    doc = Document()

    # Add title heading
    title_para = doc.add_heading('Executive Summary', level=0)

    # First paragraph - plain text, no background, no special formatting
    para1_text = ('Acme Corporation achieved record revenue of $45.2 million in fiscal year 2024, '
                  'representing a 15% year-over-year growth. This performance exceeded analyst '
                  'expectations by 3.2 percentage points.')
    para1 = doc.add_paragraph()
    run1 = para1.add_run(para1_text)
    run1.font.name = 'Calibri'
    run1.font.size = Pt(12)
    run1.bold = False

    # Second paragraph - plain text, no background, no special formatting
    para2_text = ('Operating expenses were reduced to $31.8 million, a savings of $2.4 million '
                  'compared to the previous year.')
    para2 = doc.add_paragraph()
    run2 = para2.add_run(para2_text)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(12)
    run2.bold = False

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
