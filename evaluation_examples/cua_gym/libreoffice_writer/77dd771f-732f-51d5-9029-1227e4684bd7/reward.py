"""
Reward Script: Custom list style for procedure manual
Task ID: writer_rd_060
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Level 1 items use 'Step X:' format, bold, blue (#0000CC), indented ~1.0cm
  Component 2 (0.30): Level 2 items use 'a)', 'b)', 'c)' format, indented ~2.0cm
  Component 3 (0.20): Level 3 items use bullet character, indented ~3.0cm
  Component 4 (0.15): All 3 procedures have the hierarchical formatting applied
"""

import os
import re
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_060'


def color_distance_rgb(c1, c2):
    """Euclidean distance between two RGB tuples."""
    return sqrt(sum((a - b) ** 2 for a, b in zip(c1, c2)))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Emu, Cm
    except ImportError as e:
        print(f"CRITICAL: Cannot import docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Classify paragraphs by their role in the document
    # We need to find Level 1 (step), Level 2 (sub-item), Level 3 (bullet) paragraphs
    # across the 3 procedures.

    # Identify procedure sections by Heading 1 paragraphs
    procedure_indices = []
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == 'Heading 1':
            procedure_indices.append(i)

    if len(procedure_indices) < 3:
        print(f"WARN: Expected 3 procedure headings, found {len(procedure_indices)}")

    # For each paragraph, determine if it's a Level 1, Level 2, or Level 3 item
    # based on text patterns in the golden version:
    # Level 1: starts with "Step N:"
    # Level 2: starts with a lowercase letter followed by ")" like "a)", "b)"
    # Level 3: starts with bullet character or similar bullet

    level1_paras = []  # (index, para)
    level2_paras = []
    level3_paras = []

    step_pattern = re.compile(r'^Step\s+\d+\s*:', re.IGNORECASE)
    subletter_pattern = re.compile(r'^[a-z]\)\s')
    bullet_pattern = re.compile(r'^\u2022\s')  # Unicode bullet

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if step_pattern.match(text):
            level1_paras.append((i, para))
        elif subletter_pattern.match(text):
            level2_paras.append((i, para))
        elif bullet_pattern.match(text):
            level3_paras.append((i, para))

    print(f"INFO: Found {len(level1_paras)} Level 1 (Step), {len(level2_paras)} Level 2 (a/b/c), {len(level3_paras)} Level 3 (bullet) paragraphs")

    # -------------------------------------------------------------------------
    # Component 1: Level 1 items — 'Step X:' format, bold, blue, ~1.0cm indent (0.35 pts)
    # -------------------------------------------------------------------------
    try:
        if len(level1_paras) == 0:
            print("FAIL: Component 1 — No 'Step X:' paragraphs found")
        else:
            bold_blue_count = 0
            indent_ok_count = 0
            target_blue = (0x00, 0x00, 0xCC)

            for idx, para in level1_paras:
                # Check bold and blue on runs
                has_bold_blue = False
                for run in para.runs:
                    if run.text.strip():
                        is_bold = run.font.bold is True
                        is_blue = False
                        if run.font.color and run.font.color.rgb:
                            rgb = run.font.color.rgb
                            dist = color_distance_rgb((rgb[0], rgb[1], rgb[2]), target_blue)
                            is_blue = dist < 30  # tolerance
                        if is_bold and is_blue:
                            has_bold_blue = True
                            break
                if has_bold_blue:
                    bold_blue_count += 1

                # Check indent (~1.0cm = 360000 EMU, tolerance +/- 0.3cm)
                left_indent = para.paragraph_format.left_indent
                if left_indent is not None:
                    indent_cm = left_indent / 360000.0
                    if 0.5 <= indent_cm <= 1.5:
                        indent_ok_count += 1

            total_l1 = len(level1_paras)
            bold_blue_ratio = bold_blue_count / total_l1
            indent_ratio = indent_ok_count / total_l1

            # Need both bold+blue AND indent for full credit
            # At least 80% of Level 1 items must pass each check
            sub1 = 0.0
            if bold_blue_ratio >= 0.8:
                sub1 += 0.20
                print(f"PASS: Component 1a — {bold_blue_count}/{total_l1} Level 1 items are bold blue (0.20 pts)")
            else:
                print(f"FAIL: Component 1a — Only {bold_blue_count}/{total_l1} Level 1 items are bold blue")

            if indent_ratio >= 0.8:
                sub1 += 0.15
                print(f"PASS: Component 1b — {indent_ok_count}/{total_l1} Level 1 items have ~1.0cm indent (0.15 pts)")
            else:
                print(f"FAIL: Component 1b — Only {indent_ok_count}/{total_l1} Level 1 items have ~1.0cm indent")

            total_score += sub1
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Level 2 items — 'a)', 'b)', 'c)' format, ~2.0cm indent (0.30 pts)
    # -------------------------------------------------------------------------
    try:
        if len(level2_paras) == 0:
            print("FAIL: Component 2 — No 'a)', 'b)' paragraphs found")
        else:
            indent_ok_count = 0
            format_ok_count = 0

            for idx, para in level2_paras:
                text = para.text.strip()
                # Check format: starts with letter + ")"
                if subletter_pattern.match(text):
                    format_ok_count += 1

                # Check indent (~2.0cm, tolerance +/- 0.5cm)
                left_indent = para.paragraph_format.left_indent
                if left_indent is not None:
                    indent_cm = left_indent / 360000.0
                    if 1.3 <= indent_cm <= 2.7:
                        indent_ok_count += 1

            total_l2 = len(level2_paras)
            format_ratio = format_ok_count / total_l2
            indent_ratio = indent_ok_count / total_l2

            sub2 = 0.0
            if format_ratio >= 0.8:
                sub2 += 0.15
                print(f"PASS: Component 2a — {format_ok_count}/{total_l2} Level 2 items use 'x)' format (0.15 pts)")
            else:
                print(f"FAIL: Component 2a — Only {format_ok_count}/{total_l2} Level 2 items use 'x)' format")

            if indent_ratio >= 0.8:
                sub2 += 0.15
                print(f"PASS: Component 2b — {indent_ok_count}/{total_l2} Level 2 items have ~2.0cm indent (0.15 pts)")
            else:
                print(f"FAIL: Component 2b — Only {indent_ok_count}/{total_l2} Level 2 items have ~2.0cm indent")

            total_score += sub2
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Level 3 items — bullet character (U+2022), ~3.0cm indent (0.20 pts)
    # -------------------------------------------------------------------------
    try:
        if len(level3_paras) == 0:
            print("FAIL: Component 3 — No bullet (U+2022) paragraphs found")
        else:
            indent_ok_count = 0

            for idx, para in level3_paras:
                # Check indent (~3.0cm, tolerance +/- 0.5cm)
                left_indent = para.paragraph_format.left_indent
                if left_indent is not None:
                    indent_cm = left_indent / 360000.0
                    if 2.3 <= indent_cm <= 3.7:
                        indent_ok_count += 1

            total_l3 = len(level3_paras)
            indent_ratio = indent_ok_count / total_l3

            sub3 = 0.0
            # The bullet character itself is verified by the pattern match that classified them
            if total_l3 >= 5:
                sub3 += 0.10
                print(f"PASS: Component 3a — {total_l3} Level 3 items use bullet character (0.10 pts)")
            else:
                print(f"FAIL: Component 3a — Only {total_l3} Level 3 items found (expected >= 5)")

            if indent_ratio >= 0.8:
                sub3 += 0.10
                print(f"PASS: Component 3b — {indent_ok_count}/{total_l3} Level 3 items have ~3.0cm indent (0.10 pts)")
            else:
                print(f"FAIL: Component 3b — Only {indent_ok_count}/{total_l3} Level 3 items have ~3.0cm indent")

            total_score += sub3
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: All 3 procedures have hierarchical formatting (0.15 pts)
    # Each procedure must have at least one Level 1 and one Level 2 item
    # -------------------------------------------------------------------------
    try:
        # Find procedure boundaries
        proc_starts = []
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name == 'Heading 1':
                proc_starts.append(i)

        # Add end sentinel
        proc_starts.append(len(doc.paragraphs))

        procedures_with_hierarchy = 0
        for p in range(len(proc_starts) - 1):
            start = proc_starts[p]
            end = proc_starts[p + 1]

            has_l1 = False
            has_l2 = False
            for idx, para in level1_paras:
                if start < idx < end:
                    has_l1 = True
                    break
            for idx, para in level2_paras:
                if start < idx < end:
                    has_l2 = True
                    break

            if has_l1 and has_l2:
                procedures_with_hierarchy += 1

        if procedures_with_hierarchy >= 3:
            total_score += 0.15
            print(f"PASS: Component 4 — All {procedures_with_hierarchy} procedures have hierarchical formatting (0.15 pts)")
        else:
            print(f"FAIL: Component 4 — Only {procedures_with_hierarchy}/3 procedures have hierarchical formatting")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
