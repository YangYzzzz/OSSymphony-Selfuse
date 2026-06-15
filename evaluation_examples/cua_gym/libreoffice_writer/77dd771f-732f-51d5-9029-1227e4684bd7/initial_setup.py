"""
Initial Setup: Create a procedure manual with plain numbered paragraphs (no custom list styles)
Task ID: writer_rd_060
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_060'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Document title
    title = doc.add_heading('Technical Procedure Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        'This manual outlines the standard operating procedures for the '
        'Network Operations Center (NOC). All technicians must follow these '
        'procedures to ensure consistent service delivery and minimize downtime.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # ============================================================
    # Procedure 1: Server Health Check
    # ============================================================
    doc.add_heading('Procedure 1: Server Health Check', level=1)
    doc.add_paragraph(
        'This procedure must be performed at the start of every shift to verify '
        'that all production servers are operating within normal parameters.'
    )

    # Level 1 steps as plain numbered paragraphs
    doc.add_paragraph('1. Log into the monitoring dashboard at https://monitor.internal.corp')
    doc.add_paragraph('2. Review the alert queue for any critical or warning notifications')
    doc.add_paragraph('3. Check CPU and memory utilization across all server clusters')

    # Sub-steps as plain indented paragraphs (no list formatting)
    p = doc.add_paragraph('   a. Compare current values against the 7-day rolling average')
    p = doc.add_paragraph('   b. Flag any server exceeding 85% utilization for more than 10 minutes')
    p = doc.add_paragraph('   c. Document anomalies in the shift handover log')

    doc.add_paragraph('4. Verify disk space usage on database and file servers')

    p = doc.add_paragraph('   a. Check primary storage arrays for capacity warnings')
    p = doc.add_paragraph('   b. Review automated cleanup job logs from the previous 24 hours')

    # Some deeper sub-items as plain text
    p = doc.add_paragraph('      - Confirm temp file purge completed successfully')
    p = doc.add_paragraph('      - Verify log rotation ran without errors')
    p = doc.add_paragraph('      - Check archived data transfer to cold storage')

    doc.add_paragraph('5. Run the automated health check script on all production nodes')
    doc.add_paragraph('6. Record results in the daily operations spreadsheet')

    # ============================================================
    # Procedure 2: Network Incident Response
    # ============================================================
    doc.add_heading('Procedure 2: Network Incident Response', level=1)
    doc.add_paragraph(
        'Follow this procedure when a network connectivity issue is detected '
        'or reported by end users. Timely response is critical to maintaining SLA compliance.'
    )

    doc.add_paragraph('1. Acknowledge the incident in the ticketing system within 5 minutes')
    doc.add_paragraph('2. Classify severity based on the impact matrix')

    p = doc.add_paragraph('   a. Severity 1: Complete loss of service affecting 50+ users')
    p = doc.add_paragraph('   b. Severity 2: Degraded performance affecting a department')
    p = doc.add_paragraph('   c. Severity 3: Isolated issue affecting fewer than 5 users')

    doc.add_paragraph('3. Perform initial diagnostics on affected network segments')

    p = doc.add_paragraph('   a. Run traceroute to isolate the failure point')
    p = doc.add_paragraph('   b. Check switch and router logs for error messages')

    p = doc.add_paragraph('      - Review spanning tree topology changes')
    p = doc.add_paragraph('      - Check for duplicate IP address conflicts')
    p = doc.add_paragraph('      - Verify VLAN configuration consistency')

    p = doc.add_paragraph('   c. Test physical layer connectivity where applicable')

    doc.add_paragraph('4. Escalate to the network engineering team if not resolved within 30 minutes')
    doc.add_paragraph('5. Update the incident ticket with findings every 15 minutes')
    doc.add_paragraph('6. Conduct a post-incident review within 48 hours of resolution')

    # ============================================================
    # Procedure 3: Backup Verification and Recovery Testing
    # ============================================================
    doc.add_heading('Procedure 3: Backup Verification and Recovery Testing', level=1)
    doc.add_paragraph(
        'This monthly procedure ensures that backup systems are functioning correctly '
        'and that data can be recovered within the defined Recovery Time Objective (RTO).'
    )

    doc.add_paragraph('1. Select three random production databases for recovery testing')
    doc.add_paragraph('2. Verify the latest backup integrity using checksums')

    p = doc.add_paragraph('   a. Compare SHA-256 hashes of backup files against the manifest')
    p = doc.add_paragraph('   b. Confirm backup timestamps match the scheduled backup window')

    p = doc.add_paragraph('      - Check full backup completed before 03:00 UTC')
    p = doc.add_paragraph('      - Verify incremental backups ran at 06:00, 12:00, and 18:00 UTC')

    doc.add_paragraph('3. Perform test restoration to the staging environment')

    p = doc.add_paragraph('   a. Restore the selected databases to the staging cluster')
    p = doc.add_paragraph('   b. Run data validation queries against restored data')
    p = doc.add_paragraph('   c. Measure actual recovery time and compare against RTO')

    p = doc.add_paragraph('      - RTO target for Tier 1 databases: 2 hours')
    p = doc.add_paragraph('      - RTO target for Tier 2 databases: 8 hours')
    p = doc.add_paragraph('      - RTO target for Tier 3 databases: 24 hours')

    doc.add_paragraph('4. Document results in the monthly backup compliance report')
    doc.add_paragraph('5. Report any failures to the IT Security and Compliance team')

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
