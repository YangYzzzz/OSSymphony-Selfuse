"""
Initial Setup: Demote heading 'Implementation Details' from Heading 1 to Heading 3
Task ID: writer_struct_078
Domain: libreoffice_writer

Creates a 9-page software design document with realistic content.
'Implementation Details' is Heading 1 in the initial state (before the task).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user/Desktop'
TASK_ID = 'software_design'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # ---- Title ----
    doc.add_heading('Software Design Document', level=0)
    p = doc.add_paragraph(
        'Project: Nexus Analytics Platform  |  Version: 2.4  |  Status: Draft\n'
        'Author: Engineering Team  |  Date: 2025-03-10'
    )
    p.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # ---- H1: Overview ----
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(
        'The Nexus Analytics Platform (NAP) is a cloud-native solution designed to '
        'process, store, and visualize large-scale business intelligence data. '
        'This document describes the system architecture, requirements, design '
        'decisions, and implementation strategy for the 2.4 release cycle.'
    )
    doc.add_paragraph(
        'NAP is intended for enterprise customers who need real-time dashboards, '
        'scheduled reports, and ad-hoc query capabilities against petabyte-scale '
        'data warehouses. The platform replaces the legacy BatchReport system '
        'that has been in production since 2017.'
    )

    doc.add_heading('Scope', level=2)
    doc.add_paragraph(
        'This release focuses on three key areas: (1) migrating the ingestion '
        'pipeline from Apache Kafka to Apache Pulsar, (2) introducing a columnar '
        'storage backend based on Apache Parquet, and (3) delivering a new '
        'React-based dashboard UI to replace the aging AngularJS frontend.'
    )

    doc.add_heading('Audience', level=2)
    doc.add_paragraph(
        'This document is intended for software engineers, solution architects, '
        'QA engineers, and technical project managers involved in the NAP 2.4 '
        'development cycle.'
    )

    doc.add_page_break()

    # ---- H1: Requirements ----
    doc.add_heading('Requirements', level=1)
    doc.add_paragraph(
        'This section enumerates the functional and non-functional requirements '
        'gathered during the discovery phase in Q4 2024.'
    )

    doc.add_heading('Functional Requirements', level=2)
    reqs_functional = [
        'FR-001: The system SHALL ingest data from at least 50 concurrent producers at 100 MB/s aggregate throughput.',
        'FR-002: The system SHALL support SQL-compatible query language with full ANSI-2011 compliance.',
        'FR-003: Dashboards SHALL refresh with sub-3-second latency for datasets up to 10 million rows.',
        'FR-004: The platform SHALL export reports in PDF, CSV, and Excel formats.',
        'FR-005: Role-based access control SHALL be enforced at row and column granularity.',
        'FR-006: Scheduled reports SHALL support cron-style scheduling with timezone awareness.',
        'FR-007: The system SHALL provide a REST API for all data access operations.',
        'FR-008: Audit logs SHALL capture every query and administrative action with user attribution.',
    ]
    for req in reqs_functional:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('Non-Functional Requirements', level=2)
    reqs_nonfunctional = [
        'NFR-001: System availability SHALL be 99.95% measured monthly excluding planned maintenance.',
        'NFR-002: Mean time to recovery (MTTR) SHALL be under 15 minutes for P1 incidents.',
        'NFR-003: All data at rest SHALL be encrypted using AES-256-GCM.',
        'NFR-004: All inter-service communication SHALL use mutual TLS (mTLS).',
        'NFR-005: The platform SHALL scale horizontally to support 10x current peak load without manual intervention.',
        'NFR-006: Query response p99 latency SHALL remain below 5 seconds for any adhoc query.',
    ]
    for req in reqs_nonfunctional:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('Constraints', level=2)
    doc.add_paragraph(
        'The platform must operate within existing Kubernetes-based infrastructure '
        'on AWS us-east-1. All third-party dependencies must have active community '
        'support and permissive open-source licenses (Apache 2.0 or MIT preferred). '
        'Deployment must complete within a 4-hour maintenance window on the first '
        'Sunday of each month.'
    )

    doc.add_page_break()

    # ---- H1: Design ----
    doc.add_heading('Design', level=1)
    doc.add_paragraph(
        'This section describes the high-level design decisions and architectural '
        'patterns adopted for the NAP 2.4 release. Each subsystem is described '
        'independently with clear interface contracts between layers.'
    )

    # ---- H2: Technical Architecture (under Design) ----
    doc.add_heading('Technical Architecture', level=2)
    doc.add_paragraph(
        'NAP 2.4 follows a layered, microservices architecture composed of five '
        'primary tiers: (1) Ingestion Tier, (2) Processing Tier, (3) Storage Tier, '
        '(4) Query Tier, and (5) Presentation Tier. Each tier communicates over '
        'well-defined gRPC contracts versioned independently.'
    )
    doc.add_paragraph(
        'The Ingestion Tier uses Apache Pulsar topics partitioned by tenant ID. '
        'The Processing Tier contains Apache Flink jobs that perform schema '
        'validation, deduplication, and enrichment before forwarding records to '
        'the Storage Tier. The Storage Tier persists data in Parquet format on '
        'Amazon S3 with a Delta Lake transaction log for ACID semantics.'
    )

    doc.add_heading('Service Mesh', level=3)
    doc.add_paragraph(
        'Istio 1.20 is deployed as the service mesh layer providing mTLS, '
        'traffic management, and observability. All services emit OpenTelemetry '
        'spans that are collected by a Grafana Tempo backend and visualized '
        'through Grafana dashboards. Prometheus scrapes metrics every 15 seconds.'
    )

    doc.add_heading('Data Model', level=2)
    doc.add_paragraph(
        'The canonical data model uses a star schema with a central Fact table '
        'flanked by Dimension tables for Time, Organization, Product, and Region. '
        'Foreign keys use surrogate integer keys generated by a distributed '
        'Snowflake-style ID generator to avoid hotspots on sequential inserts.'
    )
    doc.add_paragraph(
        'Schema evolution follows Avro compatibility rules: BACKWARD_TRANSITIVE '
        'for consumer safety. Schema versions are stored in the Confluent Schema '
        'Registry (compatible API) running as an in-cluster service.'
    )

    doc.add_heading('Security Design', level=2)
    doc.add_paragraph(
        'Authentication is handled by Keycloak 24 as the OpenID Connect provider. '
        'JWT tokens carry group membership claims that are mapped to RBAC roles '
        'within NAP. Token lifetime is 15 minutes; refresh tokens expire in 8 hours. '
        'All tokens are validated at the API gateway before reaching downstream services.'
    )

    doc.add_page_break()

    doc.add_heading('API Design', level=2)
    doc.add_paragraph(
        'The public REST API follows OpenAPI 3.1 specification. Versioning uses '
        'URI path prefix: /api/v2/. Breaking changes require a new major version. '
        'All responses use JSON with application/json content type. Error responses '
        'follow RFC 7807 (Problem Details for HTTP APIs).'
    )

    doc.add_heading('Pagination', level=3)
    doc.add_paragraph(
        'Large collection endpoints use cursor-based pagination via the "cursor" '
        'query parameter. Page size defaults to 50 and is capped at 500. Clients '
        'should check the "next_cursor" field in the response envelope to determine '
        'if more records are available.'
    )

    doc.add_heading('Rate Limiting', level=3)
    doc.add_paragraph(
        'Each API key is subject to a token-bucket rate limiter with a burst '
        'capacity of 200 requests and a sustained rate of 100 requests per second. '
        'Responses include X-RateLimit-Remaining and X-RateLimit-Reset headers.'
    )

    doc.add_page_break()

    # ---- H1: Implementation Details (THIS IS THE KEY HEADING — MUST BE Heading 1 in initial) ----
    doc.add_heading('Implementation Details', level=1)
    doc.add_paragraph(
        'This section provides engineering-level details covering technology stack '
        'selection, deployment procedures, and operational considerations for '
        'the NAP 2.4 release.'
    )

    doc.add_heading('Technology Stack', level=2)
    stack_items = [
        'Language: Python 3.12 (backend services), TypeScript 5.3 (frontend)',
        'Framework: FastAPI 0.111 (REST API), React 18 + Vite 5 (frontend)',
        'Message Broker: Apache Pulsar 3.2.3',
        'Stream Processing: Apache Flink 1.18',
        'Object Storage: Amazon S3 (ap-southeast-1)',
        'Table Format: Delta Lake 3.1 on Parquet',
        'Query Engine: Trino 440',
        'Cache: Redis 7.2 (Cluster mode)',
        'Identity Provider: Keycloak 24',
        'Container Runtime: containerd 1.7 on Kubernetes 1.30',
        'Service Mesh: Istio 1.20',
        'Observability: OpenTelemetry + Grafana Stack (Loki, Tempo, Prometheus)',
    ]
    for item in stack_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Deployment Pipeline', level=2)
    doc.add_paragraph(
        'CI/CD is managed via GitHub Actions with environment promotion through '
        'dev → staging → production. All Docker images are built using multi-stage '
        'Dockerfiles and pushed to Amazon ECR. Kubernetes manifests are templated '
        'with Helm 3 and deployed via ArgoCD using GitOps principles. Production '
        'deployments require two approvals from the Platform Engineering team.'
    )

    doc.add_heading('Database Migrations', level=3)
    doc.add_paragraph(
        'Schema changes to the metadata PostgreSQL database use Flyway migration '
        'scripts stored in the db/migrations/ directory. Migration scripts are '
        'automatically applied during the deployment pipeline after passing a '
        'dry-run check against a staging replica.'
    )

    doc.add_heading('Rollback Procedures', level=3)
    doc.add_paragraph(
        'ArgoCD maintains a 10-revision deployment history. Rollback to a prior '
        'version requires triggering the rollback-production GitHub Actions workflow '
        'with the target revision SHA. Automated canary analysis runs for 10 minutes '
        'before full rollout; a failure automatically triggers rollback.'
    )

    doc.add_page_break()

    doc.add_heading('Monitoring and Alerting', level=2)
    doc.add_paragraph(
        'All services expose a /metrics endpoint scraped by Prometheus every 15 '
        'seconds. Alerting rules are defined in prometheus-alerts.yaml with '
        'notification routing to PagerDuty for P1 and P2 incidents, and to '
        'Slack #alerts-engineering for P3 and below.'
    )

    doc.add_heading('SLO Dashboard', level=3)
    doc.add_paragraph(
        'The SLO dashboard in Grafana tracks availability, latency p50/p95/p99, '
        'and error rate per service. Monthly SLO reviews are conducted during the '
        'Platform Engineering sync. Error budget burn-rate alerts fire when the '
        '1-hour burn rate exceeds 14.4× the acceptable rate.'
    )

    doc.add_page_break()

    doc.add_heading('Testing Strategy', level=2)
    doc.add_paragraph(
        'The test pyramid follows a 70/20/10 split: 70% unit tests, 20% integration '
        'tests, 10% end-to-end tests. Unit tests use pytest with hypothesis for '
        'property-based testing. Integration tests spin up dependent services via '
        'Testcontainers. E2E tests use Playwright against the full staging environment.'
    )
    doc.add_paragraph(
        'Performance benchmarks run nightly using k6 load profiles that simulate '
        '1,000 concurrent users. Benchmark results are posted as GitHub PR comments '
        'comparing HEAD against the main branch baseline.'
    )

    doc.add_heading('Test Coverage Requirements', level=3)
    doc.add_paragraph(
        'All new Python modules must achieve minimum 85% line coverage. Coverage '
        'reports are uploaded to Codecov and a PR check fails if coverage drops '
        'below the threshold for the changed files. Frontend components must have '
        'at least one snapshot test and one interaction test per component.'
    )

    doc.add_page_break()

    # ---- H1: Appendix ----
    doc.add_heading('Appendix', level=1)

    doc.add_heading('Glossary', level=2)
    glossary_items = [
        'ACID — Atomicity, Consistency, Isolation, Durability',
        'CRD — Custom Resource Definition (Kubernetes)',
        'GitOps — Infrastructure-as-code practice where Git is the source of truth',
        'mTLS — Mutual Transport Layer Security',
        'RBAC — Role-Based Access Control',
        'SLO — Service Level Objective',
        'WAL — Write-Ahead Log',
    ]
    for item in glossary_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('References', level=2)
    refs = [
        'Apache Pulsar Documentation 3.2 — https://pulsar.apache.org/docs/3.2.x/',
        'Delta Lake Protocol Specification — https://github.com/delta-io/delta/blob/master/PROTOCOL.md',
        'OpenTelemetry Specification 1.30 — https://opentelemetry.io/docs/specs/',
        'RFC 7807 — Problem Details for HTTP APIs — https://www.rfc-editor.org/rfc/rfc7807',
        'Kubernetes Documentation 1.30 — https://kubernetes.io/docs/home/',
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='List Bullet')

    doc.add_heading('Revision History', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Version'
    hdr[1].text = 'Date'
    hdr[2].text = 'Author'
    hdr[3].text = 'Summary'
    revisions = [
        ('2.4', '2025-03-10', 'Diana Park', 'Initial draft for 2.4 release'),
        ('2.3', '2024-11-22', 'Marcus Liu', 'Added security design section'),
        ('2.2', '2024-08-05', 'Priya Sharma', 'Revised non-functional requirements'),
        ('2.1', '2024-04-18', 'James O\'Brien', 'Data model updates for star schema'),
        ('2.0', '2024-01-09', 'Diana Park', 'Major rewrite for microservices migration'),
    ]
    for rev in revisions:
        row = table.add_row().cells
        for i, val in enumerate(rev):
            row[i].text = val

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
