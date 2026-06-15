"""
Reward Script: Demote 'Implementation Details' from Heading 1 to Heading 3
Task ID: writer_struct_078
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): 'Implementation Details' paragraph has style 'Heading 3'
  Component 2 (0.4): 'Implementation Details' is NOT a Heading 1 AND
                     the 4 original Heading 1 entries remain intact
"""

import os
from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_struct_078'
FILE_PATH = f'{WORKDIR}/software_design.docx'

# Expected Heading 1 entries after the task (Implementation Details removed)
EXPECTED_H1_TEXTS = {'Overview', 'Requirements', 'Design', 'Appendix'}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: 'Implementation Details' has style 'Heading 3' (0.6 points)
    # On initial_env: style is 'Heading 1' -> FAILS -> 0 pts
    # On golden_env:  style is 'Heading 3' -> PASSES -> 0.6 pts
    try:
        impl_details_style = None
        for para in doc.paragraphs:
            if para.text.strip() == 'Implementation Details':
                impl_details_style = para.style.name
                break

        if impl_details_style is None:
            print("FAIL: Component 1 — paragraph 'Implementation Details' not found in document")
        elif impl_details_style == 'Heading 3':
            print(f"PASS: Component 1 — 'Implementation Details' has style 'Heading 3' (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 1 — 'Implementation Details' has style '{impl_details_style}', expected 'Heading 3'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 'Implementation Details' is NOT a Heading 1, and the original Heading 1
    # entries (Overview, Requirements, Design, Appendix) remain intact (0.4 points)
    # On initial_env: Implementation Details IS Heading 1 -> FAILS -> 0 pts
    # On golden_env:  Implementation Details is Heading 3 and only 4 H1 entries remain -> PASSES -> 0.4 pts
    try:
        actual_h1_texts = set(
            para.text.strip()
            for para in doc.paragraphs
            if para.style.name == 'Heading 1'
        )
        impl_details_in_h1 = 'Implementation Details' in actual_h1_texts
        original_h1_intact = EXPECTED_H1_TEXTS.issubset(actual_h1_texts)

        if impl_details_in_h1:
            print(f"FAIL: Component 2 — 'Implementation Details' is still a Heading 1")
        elif not original_h1_intact:
            missing = EXPECTED_H1_TEXTS - actual_h1_texts
            print(f"FAIL: Component 2 — Missing expected Heading 1 entries: {missing}")
        elif not impl_details_in_h1 and original_h1_intact:
            print(f"PASS: Component 2 — 'Implementation Details' removed from Heading 1; "
                  f"original H1 entries intact: {sorted(actual_h1_texts)} (0.4 pts)")
            total_score += 0.4
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
