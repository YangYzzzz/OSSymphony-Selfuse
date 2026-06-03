"""
Initial Setup: Journal entry document with plain text (no heading styles applied)
Task ID: writer_creative_042
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'journal_2026'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Remove default paragraph style excess spacing from Normal style
    # All paragraphs should use Default Paragraph Style with 12pt font

    # --- Document Title (plain Default Paragraph Style, NOT Heading 1) ---
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('My Journal \u2014 2026')
    title_run.font.size = Pt(12)

    # --- Entry 1: January 15, 2026 ---
    date_para1 = doc.add_paragraph()
    date_run1 = date_para1.add_run('January 15, 2026')
    date_run1.font.size = Pt(12)

    body_para1a = doc.add_paragraph()
    body_para1a.add_run(
        'Started the new year with a resolution to write more consistently. '
        'The morning was cold but clear, and I sat by the window watching the frost '
        'patterns on the glass while drinking my coffee.'
    ).font.size = Pt(12)

    body_para1b = doc.add_paragraph()
    body_para1b.add_run(
        'Called my sister Elena in the evening. She mentioned she might visit in March, '
        'which would be wonderful. We talked for almost two hours about everything and nothing.'
    ).font.size = Pt(12)

    # --- Entry 2: February 3, 2026 ---
    date_para2 = doc.add_paragraph()
    date_run2 = date_para2.add_run('February 3, 2026')
    date_run2.font.size = Pt(12)

    body_para2a = doc.add_paragraph()
    body_para2a.add_run(
        'The project at work finally got approved after months of back-and-forth with the '
        'stakeholders. Our team celebrated with lunch at that new Thai place on Marchmont Street.'
    ).font.size = Pt(12)

    body_para2b = doc.add_paragraph()
    body_para2b.add_run(
        'I have been reading "The Mezzanine" by Nicholson Baker. '
        'It is an oddly compelling book about the minutiae of everyday life. '
        'Picked up some groceries on the way home: lemons, olive oil, dried pasta.'
    ).font.size = Pt(12)

    body_para2c = doc.add_paragraph()
    body_para2c.add_run(
        'Temperature dropped again overnight. The forecast says we might get snow by the weekend. '
        'Dug out my heavy coat from the back of the wardrobe just in case.'
    ).font.size = Pt(12)

    # --- Entry 3: February 18, 2026 ---
    date_para3 = doc.add_paragraph()
    date_run3 = date_para3.add_run('February 18, 2026')
    date_run3.font.size = Pt(12)

    body_para3a = doc.add_paragraph()
    body_para3a.add_run(
        'Spent the afternoon at the botanical gardens. Despite the grey skies, '
        'there were already early crocuses pushing through near the south-facing wall. '
        'A robin followed me along the gravel path for a good ten minutes.'
    ).font.size = Pt(12)

    body_para3b = doc.add_paragraph()
    body_para3b.add_run(
        'Finished a watercolour sketch I had been putting off for weeks. '
        'Not my best work but I am pleased I completed it. '
        'Cooking: made a big pot of lentil soup to last through the week.'
    ).font.size = Pt(12)

    # --- Entry 4: March 1, 2026 ---
    date_para4 = doc.add_paragraph()
    date_run4 = date_para4.add_run('March 1, 2026')
    date_run4.font.size = Pt(12)

    body_para4a = doc.add_paragraph()
    body_para4a.add_run(
        'First day of March and it finally feels like winter is loosening its grip. '
        'Walked to work instead of taking the bus. The cherry tree on Kendal Avenue '
        'has the first blush of blossom on its lower branches. Counted seven buds.'
    ).font.size = Pt(12)

    # --- Entry 5: March 4, 2026 ---
    date_para5 = doc.add_paragraph()
    date_run5 = date_para5.add_run('March 4, 2026')
    date_run5.font.size = Pt(12)

    body_para5a = doc.add_paragraph()
    body_para5a.add_run(
        'Elena arrived this afternoon, two days earlier than expected. '
        'Her train was delayed by forty minutes so I waited at the station with a coffee '
        'and re-read old notes on my phone. We had dinner at home: roast chicken and roasted '
        'vegetables, simple but good.'
    ).font.size = Pt(12)

    body_para5b = doc.add_paragraph()
    body_para5b.add_run(
        'Stayed up late talking. She showed me photos from her trip to Lisbon last autumn. '
        'The light in those pictures was extraordinary. We made plans to visit together '
        'sometime next year if schedules allow.'
    ).font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
