"""
Initial Setup: Chemistry lab report with chemical formulas (no subscripts applied yet)
Task ID: writer_txtfmt_033
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'lab_report'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Use Times New Roman 12pt for all runs (as specified in context)
    def add_paragraph_text(doc, text):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return para

    # Paragraph 1: Introduction mentioning NaCl
    add_paragraph_text(
        doc,
        'In this experiment, we investigated the solubility of NaCl (sodium chloride) in water '
        'at varying temperatures. Sodium chloride is a common ionic compound used widely in '
        'laboratory and industrial applications.'
    )

    # Paragraph 2: Background/Methods
    add_paragraph_text(
        doc,
        'Solutions were prepared by dissolving measured quantities of each compound in 100 mL of '
        'distilled water. Temperature was controlled using a calibrated water bath, and '
        'measurements were recorded at 10-degree intervals from 10°C to 80°C.'
    )

    # Paragraph 3: Mentions H2SO4 (no subscripts in initial)
    para3 = doc.add_paragraph()
    run3a = para3.add_run(
        'In the second phase of the experiment, we examined the reaction between zinc metal and '
        'H2SO4 (sulfuric acid). The reaction produces hydrogen gas and zinc sulfate as products, '
        'which were collected and analyzed for purity.'
    )
    run3a.font.name = 'Times New Roman'
    run3a.font.size = Pt(12)

    # Paragraph 4: Results
    add_paragraph_text(
        doc,
        'Results showed a consistent increase in reaction rate with rising temperature. '
        'The activation energy was calculated using the Arrhenius equation, yielding a value '
        'of 45.3 kJ/mol, which is consistent with published literature values for this reaction.'
    )

    # Paragraph 5: Mentions Fe2O3 (no subscripts in initial)
    para5 = doc.add_paragraph()
    run5a = para5.add_run(
        'Finally, we studied the reduction of Fe2O3 (iron(III) oxide) using carbon monoxide at '
        'elevated temperatures. This reaction is fundamental to industrial iron smelting processes '
        'and the thermodynamic feasibility was confirmed by calculating the Gibbs free energy change.'
    )
    run5a.font.name = 'Times New Roman'
    run5a.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
