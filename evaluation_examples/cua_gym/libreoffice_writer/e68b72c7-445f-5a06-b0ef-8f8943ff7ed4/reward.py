"""
Reward Script: Format chemical formula subscripts in a chemistry lab report
Task ID: writer_txtfmt_033
Domain: libreoffice_writer
Scoring:
  - Component 1a: '2' in H2SO4 (paragraph 3, 0-indexed 2) is subscripted   — 0.25 pts
  - Component 1b: '4' in H2SO4 (paragraph 3, 0-indexed 2) is subscripted   — 0.25 pts
  - Component 2a: '2' in Fe2O3 (paragraph 5, 0-indexed 4) is subscripted   — 0.25 pts
  - Component 2b: '3' in Fe2O3 (paragraph 5, 0-indexed 4) is subscripted   — 0.25 pts
  Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user'
FILE_PATH = os.path.join(WORKDIR, 'Desktop', 'lab_report.docx')


def get_subscripted_chars(para):
    """
    Returns a set of text strings from runs that have explicit subscript=True
    in the given paragraph. Checks character-by-character if runs contain
    multiple characters.
    """
    subscript_set = set()
    for run in para.runs:
        if run.font.subscript is True:
            # Add each individual character from this subscripted run
            for ch in run.text:
                subscript_set.add(ch)
    return subscript_set


def verify_task(file_path):
    """
    Verify that chemical formula subscripts have been applied correctly:
    - In paragraph 3 (1-indexed) / index 2: H2SO4 must have '2' and '4' subscripted
    - In paragraph 5 (1-indexed) / index 4: Fe2O3 must have '2' and '3' subscripted

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 5 paragraphs
    if len(doc.paragraphs) < 5:
        print(f"CRITICAL: Expected at least 5 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Paragraph indices (0-based):
    #   Task paragraph 3 (1-based) = index 2 — contains H2SO4
    #   Task paragraph 5 (1-based) = index 4 — contains Fe2O3
    para_h2so4 = doc.paragraphs[2]
    para_fe2o3 = doc.paragraphs[4]

    # ------------------------------------------------------------------
    # Component 1a: '2' in H2SO4 (paragraph index 2) is subscripted (0.25 pts)
    # Task change: '2' goes from regular text to subscript format
    # ------------------------------------------------------------------
    try:
        subscripted_h2so4 = get_subscripted_chars(para_h2so4)
        if '2' in subscripted_h2so4:
            print("PASS: Component 1a — '2' in H2SO4 (para 3) is subscripted (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1a — '2' in H2SO4 is NOT subscripted. "
                  f"Subscripted chars found: {subscripted_h2so4}")
    except Exception as e:
        print(f"ERROR: Component 1a — {e}")

    # ------------------------------------------------------------------
    # Component 1b: '4' in H2SO4 (paragraph index 2) is subscripted (0.25 pts)
    # Task change: '4' goes from regular text to subscript format
    # ------------------------------------------------------------------
    try:
        subscripted_h2so4 = get_subscripted_chars(para_h2so4)
        if '4' in subscripted_h2so4:
            print("PASS: Component 1b — '4' in H2SO4 (para 3) is subscripted (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1b — '4' in H2SO4 is NOT subscripted. "
                  f"Subscripted chars found: {subscripted_h2so4}")
    except Exception as e:
        print(f"ERROR: Component 1b — {e}")

    # ------------------------------------------------------------------
    # Component 2a: '2' in Fe2O3 (paragraph index 4) is subscripted (0.25 pts)
    # Task change: '2' goes from regular text to subscript format
    # ------------------------------------------------------------------
    try:
        subscripted_fe2o3 = get_subscripted_chars(para_fe2o3)
        if '2' in subscripted_fe2o3:
            print("PASS: Component 2a — '2' in Fe2O3 (para 5) is subscripted (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2a — '2' in Fe2O3 is NOT subscripted. "
                  f"Subscripted chars found: {subscripted_fe2o3}")
    except Exception as e:
        print(f"ERROR: Component 2a — {e}")

    # ------------------------------------------------------------------
    # Component 2b: '3' in Fe2O3 (paragraph index 4) is subscripted (0.25 pts)
    # Task change: '3' goes from regular text to subscript format
    # ------------------------------------------------------------------
    try:
        subscripted_fe2o3 = get_subscripted_chars(para_fe2o3)
        if '3' in subscripted_fe2o3:
            print("PASS: Component 2b — '3' in Fe2O3 (para 5) is subscripted (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2b — '3' in Fe2O3 is NOT subscripted. "
                  f"Subscripted chars found: {subscripted_fe2o3}")
    except Exception as e:
        print(f"ERROR: Component 2b — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path on the VM
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
