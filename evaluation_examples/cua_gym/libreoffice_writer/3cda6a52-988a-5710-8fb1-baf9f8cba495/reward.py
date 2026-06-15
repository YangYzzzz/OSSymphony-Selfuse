"""
Reward Script: Envelope layout for mailing to client
Task ID: writer_biz_056
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Document has 2+ sections (envelope page added)
  Component 2 (0.25): Envelope section has #10 dimensions (9.5 x 4.125 inches)
  Component 3 (0.25): Return address present in envelope section, left-aligned
  Component 4 (0.25): Delivery address present in envelope section, centered
"""

import os
from docx import Document
from docx.shared import Inches, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_056'

# Tolerance for page dimensions: 0.15 inches (in EMU)
DIM_TOLERANCE = int(0.15 * 914400)

# Expected #10 envelope: 9.5 x 4.125 inches
ENVELOPE_WIDTH = int(9.5 * 914400)
ENVELOPE_HEIGHT = int(4.125 * 914400)

# Expected addresses
RETURN_ADDRESS_LINES = [
    'Meridian Solutions Inc.',
    '1200 Commerce Drive, Suite 400',
    'Chicago, IL 60601',
]

DELIVERY_ADDRESS_LINES = [
    'Mr. James Wilson',
    'Pinnacle Corp',
    '500 Park Avenue',
    'New York, NY 10022',
]


def persist_app_state(domain):
    """Save any unsaved LibreOffice Writer state."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def normalize(text):
    """Normalize text for comparison: strip whitespace, lowercase."""
    return ' '.join(text.strip().lower().split())


def find_envelope_section(doc):
    """
    Find the section that has envelope-like dimensions.
    Returns (section_index, section) or (None, None).
    """
    for i, sec in enumerate(doc.sections):
        w = sec.page_width
        h = sec.page_height
        # Check if dimensions match #10 envelope (9.5 x 4.125)
        if (abs(w - ENVELOPE_WIDTH) < DIM_TOLERANCE and
                abs(h - ENVELOPE_HEIGHT) < DIM_TOLERANCE):
            return i, sec
        # Also check rotated (4.125 x 9.5)
        if (abs(h - ENVELOPE_WIDTH) < DIM_TOLERANCE and
                abs(w - ENVELOPE_HEIGHT) < DIM_TOLERANCE):
            return i, sec
    return None, None


def get_paragraphs_in_section(doc, section_index):
    """
    Get paragraphs belonging to a specific section.
    Paragraphs before the first section break belong to section 0 (if the section
    break is encoded at the end of that section's last paragraph). For python-docx,
    sections[i] corresponds to the i-th sectPr. The paragraphs belonging to section i
    are those between section i-1's break and section i's break.
    """
    from lxml import etree
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    body = doc.element.body
    all_paras = list(body.iterchildren('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'))

    # Find section break positions
    # Each paragraph that contains a sectPr (in pPr) marks the end of a section
    section_breaks = []
    for idx, p_elem in enumerate(all_paras):
        pPr = p_elem.find('w:pPr', ns)
        if pPr is not None:
            sectPr = pPr.find('w:sectPr', ns)
            if sectPr is not None:
                section_breaks.append(idx)

    # Section 0: paragraphs from start to first break (inclusive)
    # Section 1: paragraphs after first break to second break (or end)
    # etc.
    if section_index == 0:
        if section_breaks:
            end = section_breaks[0] + 1
        else:
            end = len(all_paras)
        para_elems = all_paras[:end]
    else:
        if section_index - 1 < len(section_breaks):
            start = section_breaks[section_index - 1] + 1
        else:
            start = len(all_paras)
        if section_index < len(section_breaks):
            end = section_breaks[section_index] + 1
        else:
            end = len(all_paras)
        para_elems = all_paras[start:end]

    # Map element objects back to python-docx Paragraph objects
    paras = []
    para_map = {p._element: p for p in doc.paragraphs}
    for elem in para_elems:
        if elem in para_map:
            paras.append(para_map[elem])
    return paras


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)
    print(f"INFO: Document has {num_sections} section(s), {len(doc.paragraphs)} paragraph(s)")

    # Component 1: Document has 2+ sections (envelope section added) — 0.25 points
    try:
        if num_sections >= 2:
            print(f"PASS: Component 1 — Document has {num_sections} sections (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Document has only {num_sections} section(s), expected 2+")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Envelope section has #10 dimensions (9.5 x 4.125 inches) — 0.25 points
    try:
        env_idx, env_sec = find_envelope_section(doc)
        if env_sec is not None:
            w_in = env_sec.page_width / 914400
            h_in = env_sec.page_height / 914400
            print(f"PASS: Component 2 — Envelope section {env_idx} has dimensions {w_in:.3f} x {h_in:.3f} inches (0.25 pts)")
            total_score += 0.25
        else:
            # Print all section dimensions for debugging
            for i, sec in enumerate(doc.sections):
                w_in = sec.page_width / 914400
                h_in = sec.page_height / 914400
                print(f"  Section {i}: {w_in:.3f} x {h_in:.3f} inches")
            print(f"FAIL: Component 2 — No section with #10 envelope dimensions (9.5 x 4.125 in) found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Components 3 and 4 require an envelope section
    if env_sec is None:
        print("FAIL: Components 3 & 4 skipped — no envelope section found")
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Get paragraphs in the envelope section
    try:
        env_paras = get_paragraphs_in_section(doc, env_idx)
        env_texts = [(p.text.strip(), p.paragraph_format.alignment) for p in env_paras if p.text.strip()]
        print(f"INFO: Envelope section has {len(env_paras)} paragraphs ({len(env_texts)} non-empty)")
        for t, a in env_texts:
            print(f"  '{t}' align={a}")
    except Exception as e:
        print(f"ERROR: Could not extract envelope paragraphs: {e}")
        env_texts = []

    # Component 3: Return address in envelope section, left-aligned — 0.25 points
    try:
        # Check that return address lines appear in the envelope section
        found_return = 0
        left_aligned_return = 0
        for expected_line in RETURN_ADDRESS_LINES:
            for text, align in env_texts:
                if normalize(expected_line) in normalize(text):
                    found_return += 1
                    # LEFT alignment or None (inherited, typically left)
                    if align in (WD_PARAGRAPH_ALIGNMENT.LEFT, 0):
                        left_aligned_return += 1
                    break

        if found_return >= 2 and left_aligned_return >= 2:
            print(f"PASS: Component 3 — Return address found ({found_return}/{len(RETURN_ADDRESS_LINES)} lines, {left_aligned_return} left-aligned) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 3 — Return address: {found_return}/{len(RETURN_ADDRESS_LINES)} lines found, {left_aligned_return} left-aligned")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Delivery address in envelope section, centered — 0.25 points
    try:
        found_delivery = 0
        centered_delivery = 0
        for expected_line in DELIVERY_ADDRESS_LINES:
            for text, align in env_texts:
                if normalize(expected_line) in normalize(text):
                    found_delivery += 1
                    if align == WD_PARAGRAPH_ALIGNMENT.CENTER:
                        centered_delivery += 1
                    break

        if found_delivery >= 3 and centered_delivery >= 3:
            print(f"PASS: Component 4 — Delivery address found ({found_delivery}/{len(DELIVERY_ADDRESS_LINES)} lines, {centered_delivery} centered) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 4 — Delivery address: {found_delivery}/{len(DELIVERY_ADDRESS_LINES)} lines found, {centered_delivery} centered")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
