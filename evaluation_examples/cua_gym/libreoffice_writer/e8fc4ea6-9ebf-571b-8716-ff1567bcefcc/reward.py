"""
Reward Script: Insert 'Last Printed' and 'Last Saved' date fields in footer separated by pipe
Task ID: writer_tm_083
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): PRINTDATE field code exists in footer
  Component 2 (0.3): SAVEDATE field code exists in footer
  Component 3 (0.2): Pipe separator ' | ' exists in footer text between fields
  Component 4 (0.2): Both fields are proper fldChar field codes (not plain text)
"""

import os
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_083'

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': WNS}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all footer paragraphs across all sections
    footer_paragraphs = []
    for section in doc.sections:
        footer = section.footer
        if footer and footer.paragraphs:
            footer_paragraphs.extend(footer.paragraphs)

    if not footer_paragraphs:
        print("FAIL: No footer paragraphs found")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Parse all footer XML to find field codes
    has_printdate = False
    has_savedate = False
    has_pipe_separator = False
    has_proper_field_codes = False

    for para in footer_paragraphs:
        p_elem = para._element

        # Find all instrText elements to identify field codes
        instr_texts = p_elem.findall('.//w:instrText', NS)
        field_names = []
        for instr in instr_texts:
            txt = (instr.text or '').strip().upper()
            field_names.append(txt)
            if 'PRINTDATE' in txt:
                has_printdate = True
            if 'SAVEDATE' in txt:
                has_savedate = True

        # Check for fldChar elements (proper field codes, not plain text)
        fld_chars = p_elem.findall('.//w:fldChar', NS)
        begin_count = sum(1 for fc in fld_chars if fc.get(f'{{{WNS}}}fldCharType') == 'begin')
        end_count = sum(1 for fc in fld_chars if fc.get(f'{{{WNS}}}fldCharType') == 'end')
        if begin_count >= 2 and end_count >= 2:
            has_proper_field_codes = True

        # Check for pipe separator in the text
        full_text = para.text
        if ' | ' in full_text:
            has_pipe_separator = True

    # Component 1: PRINTDATE field exists in footer (0.3 points)
    try:
        if has_printdate:
            print(f"PASS: Component 1 — PRINTDATE field found in footer (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — PRINTDATE field not found in footer instrText elements")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: SAVEDATE field exists in footer (0.3 points)
    try:
        if has_savedate:
            print(f"PASS: Component 2 — SAVEDATE field found in footer (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — SAVEDATE field not found in footer instrText elements")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Pipe separator ' | ' between fields (0.2 points)
    try:
        if has_pipe_separator:
            print(f"PASS: Component 3 — Pipe separator ' | ' found in footer text (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — Pipe separator ' | ' not found in footer text")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Both fields use proper fldChar field codes (0.2 points)
    try:
        if has_proper_field_codes and has_printdate and has_savedate:
            print(f"PASS: Component 4 — Both fields use proper fldChar field codes (0.2 pts)")
            total_score += 0.2
        else:
            if not has_proper_field_codes:
                print(f"FAIL: Component 4 — Less than 2 proper field code pairs found in footer")
            else:
                print(f"FAIL: Component 4 — Missing one or both field codes")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
