"""
Initial Setup: Insert date fields in footer
Task ID: writer_tm_083
Domain: libreoffice_writer

Creates an Audit_Trail document with realistic content.
Footer is empty (no fields, no text).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_083'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title ---
    title = doc.add_heading('Quarterly Financial Audit Trail', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by Internal Audit Department')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Report Period: January 1 \u2013 March 31, 2025')
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph()  # spacer

    # --- Section 1: Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This audit trail documents all significant financial transactions and '
        'control activities performed during Q1 2025. The review encompassed '
        'accounts payable, accounts receivable, payroll processing, and general '
        'ledger reconciliations across all three regional offices.'
    )
    doc.add_paragraph(
        'Key findings indicate that 97.3% of sampled transactions were processed '
        'in compliance with established policies. Three exceptions were identified '
        'in the procurement workflow, each involving purchase orders exceeding '
        '$25,000 without the required dual-signature authorization.'
    )

    # --- Section 2: Scope and Methodology ---
    doc.add_heading('2. Scope and Methodology', level=1)
    doc.add_paragraph(
        'The audit covered 1,247 transactions totaling $8,432,019.56 across the '
        'following departments:'
    )
    items = [
        'Finance & Accounting \u2013 542 transactions ($3,891,204.10)',
        'Human Resources \u2013 318 transactions ($2,104,556.00)',
        'Operations \u2013 256 transactions ($1,672,340.88)',
        'Information Technology \u2013 131 transactions ($763,918.58)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'Statistical sampling at a 95% confidence level was applied. All source '
        'documents were cross-referenced with ERP system entries and bank '
        'reconciliation statements.'
    )

    # --- Section 3: Transaction Summary Table ---
    doc.add_heading('3. Transaction Summary', level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    headers = ['Date', 'Reference No.', 'Description', 'Amount ($)', 'Status']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data = [
        ['2025-01-08', 'AP-2025-0012', 'Vendor payment \u2013 Meridian Supplies', '14,320.00', 'Verified'],
        ['2025-01-15', 'PR-2025-0041', 'Payroll run \u2013 January cycle 1', '187,642.50', 'Verified'],
        ['2025-01-22', 'PO-2025-0087', 'Equipment purchase \u2013 Server upgrade', '32,750.00', 'Exception'],
        ['2025-02-03', 'AR-2025-0019', 'Client invoice \u2013 Vertex Solutions', '45,230.00', 'Verified'],
        ['2025-02-10', 'AP-2025-0033', 'Utility payment \u2013 Q1 electricity', '6,812.75', 'Verified'],
        ['2025-02-18', 'PO-2025-0102', 'Software license renewal \u2013 SAP', '28,900.00', 'Exception'],
        ['2025-03-01', 'PR-2025-0089', 'Payroll run \u2013 March cycle 1', '191,204.10', 'Verified'],
        ['2025-03-12', 'AP-2025-0056', 'Consulting fee \u2013 Deloitte', '55,000.00', 'Verified'],
        ['2025-03-20', 'GL-2025-0014', 'Intercompany transfer \u2013 APAC office', '120,000.00', 'Verified'],
        ['2025-03-28', 'PO-2025-0118', 'Office furniture \u2013 HQ renovation', '41,200.00', 'Exception'],
    ]
    for row_data in data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.add_paragraph()  # spacer

    # --- Section 4: Findings ---
    doc.add_heading('4. Findings and Recommendations', level=1)
    doc.add_paragraph(
        'Finding 1: Three purchase orders (PO-2025-0087, PO-2025-0102, PO-2025-0118) '
        'were approved with a single signature despite exceeding the $25,000 threshold. '
        'Recommendation: Implement automated workflow controls in the ERP system to '
        'enforce dual-signature requirements.'
    )
    doc.add_paragraph(
        'Finding 2: The intercompany transfer (GL-2025-0014) documentation was '
        'received 5 business days after the transaction date. Recommendation: '
        'Establish a 48-hour documentation submission policy for all intercompany '
        'transfers.'
    )

    # --- Section 5: Sign-off ---
    doc.add_heading('5. Audit Sign-off', level=1)
    doc.add_paragraph('Prepared by: Rachel Nguyen, CIA, Senior Internal Auditor')
    doc.add_paragraph('Reviewed by: David Kowalski, CPA, Director of Internal Audit')
    doc.add_paragraph('Date of Report: April 14, 2025')

    # --- Footer: intentionally empty ---
    footer = section.footer
    footer.is_linked_to_previous = False
    # Leave footer empty - the task is to add date fields here

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
