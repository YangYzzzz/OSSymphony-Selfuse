"""
FINAL REWARD SCRIPT - SUCCESS
Task: Standardize this document’s default font to Liberation Serif and apply it to all existing content.
Generated: 2025-10-14 12:00:46
Status: success
Model: azure-o3
Total Steps: 6
"""

from docx import Document
import zipfile
import os
from lxml import etree

"""
Reward script for the task:
"Standardize this document’s default font to Liberation Serif and apply it to all existing content."

Verification logic (progressive scoring):
1. Default/Normal style font check  – 0.5 points
   • Accepts either docDefaults or the Normal style explicitly setting all character-set
     attributes (ascii, hAnsi, eastAsia, cs) to “Liberation Serif”.
2. Existing content check           – 0.5 points
   • Inspects every run (body, tables, headers, footers).  A run is considered correct if
     its explicit font is Liberation Serif *or* it inherits from a correct default.
   • All runs must comply for the full 0.5 points.
The script prints detailed diagnostics and finally outputs:  REWARD: X.X
"""

TARGET_FONT = "Liberation Serif"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

# ------------------------------------------------------------
# Helper functions
# ------------------------------------------------------------

def _rfonts_match(rFonts, target_font=TARGET_FONT):
    """Return True if <w:rFonts> sets ascii, hAnsi, eastAsia, cs to target_font."""
    if rFonts is None:
        return False
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        val = rFonts.get(f"{{{NS['w']}}}{attr}")
        if val is None or val.strip().lower() != target_font.lower():
            return False
    return True


def check_default_font(file_path, target_font=TARGET_FONT):
    """Verify that either docDefaults or the Normal style uses the target font."""
    try:
        with zipfile.ZipFile(file_path) as z:
            styles_xml = z.read("word/styles.xml")
    except Exception as e:
        print("✗ Cannot read styles.xml:", e)
        return False

    root = etree.fromstring(styles_xml)

    # A) docDefaults
    rFonts_doc = root.find(
        ".//w:docDefaults//w:rPrDefault//w:rPr/w:rFonts", namespaces=NS
    )
    doc_ok = _rfonts_match(rFonts_doc, target_font)
    if doc_ok:
        print("✓ Document default font (docDefaults) correctly set")
    else:
        print("• docDefaults not fully set to target font (checking Normal style)")

    # B) Normal style
    normal_nodes = root.xpath("//w:style[@w:styleId='Normal']", namespaces=NS)
    normal_ok = False
    if normal_nodes:
        rFonts_norm = normal_nodes[0].find("w:rPr/w:rFonts", namespaces=NS)
        normal_ok = _rfonts_match(rFonts_norm, target_font)
        if normal_ok:
            print("✓ Normal style font correctly set")
        else:
            print("• Normal style not fully set to target font")
    else:
        print("• Normal style not found in styles.xml")

    return doc_ok or normal_ok


# -------------------------------------------------------------------------
# Paragraph/run iteration helpers (body, tables, headers, footers)
# -------------------------------------------------------------------------

def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_all_paragraphs(doc):
    # Body
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        yield from _iter_table_paragraphs(tbl)
    # Headers / footers
    for sec in doc.sections:
        for container in [
            sec.header,
            sec.first_page_header,
            sec.even_page_header,
            sec.footer,
            sec.first_page_footer,
            sec.even_page_footer,
        ]:
            if container is None:
                continue
            for p in container.paragraphs:
                yield p
            for tbl in container.tables:
                yield from _iter_table_paragraphs(tbl)


def check_runs_font(file_path, default_font_ok, target_font=TARGET_FONT):
    try:
        doc = Document(file_path)
    except Exception as e:
        print("✗ python-docx failed to open file:", e)
        return False

    total_runs = 0
    bad_runs = 0

    for para in _iter_all_paragraphs(doc):
        for run in para.runs:
            total_runs += 1
            explicit_font = run.font.name
            if explicit_font:  # run overrides default
                if explicit_font.strip().lower() != target_font.lower():
                    bad_runs += 1
            else:  # inherits from default
                if not default_font_ok:
                    bad_runs += 1

    print(f"Runs inspected: {total_runs}")
    print(f"Runs with incorrect font: {bad_runs}")

    return total_runs > 0 and bad_runs == 0


# ------------------------------------------------------------
# Main verification routine
# ------------------------------------------------------------

def verify_task(file_path):
    print("=== Verifying Liberation Serif standardization ===")
    if not os.path.exists(file_path):
        print("✗ File not found:", file_path)
        print("REWARD: 0.0")
        return 0.0

    score = 0.0

    # 1. Default / Normal style verification ---------------------------------
    default_ok = check_default_font(file_path)
    if default_ok:
        score += 0.5
        print("✓ Default font verification passed (0.5)")
    else:
        print("✗ Default font incorrect (0 points)")

    # 2. Content runs verification -------------------------------------------
    runs_ok = check_runs_font(file_path, default_ok)
    if runs_ok:
        score += 0.5
        print("✓ All existing content uses Liberation Serif (0.5)")
    else:
        print("✗ Some content still uses another font (0 points)")

    final_score = min(score, 1.0)
    print("---------------------------------------------")
    print(f"FINAL SCORE: {final_score}")
    print(f"REWARD: {final_score}")
    return final_score


if __name__ == "__main__":
    DOC_PATH = "/home/user/standardize_this_documents_default_font_to_liberation_serif_and_apply_it_to_all_existing_content.docx"
    verify_task(DOC_PATH)
