"""
Reward Script: Convert 3x4 table to tab-separated plain text
Task ID: writer_tm_011
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Table is removed from the document
  Component 2 (0.5): Four tab-separated text lines with correct content
  Component 3 (0.2): Tab-separated lines contiguous and table removed (compound check)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_011'

# Expected tab-separated rows (from the original table data)
EXPECTED_ROWS = [
    "Product\tSKU\tPrice",
    "Widget A\tW-001\t12.99",
    "Widget B\tW-002\t15.99",
    "Widget C\tW-003\t9.99",
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Pre-compute table count (used by Components 1 and 3)
    num_tables = len(doc.tables)

    # Component 1: Table is removed (0.3 points)
    # The initial document has 1 table; the golden should have 0.
    try:
        if num_tables == 0:
            print(f"PASS: Component 1 — No tables found in document (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — Expected 0 tables, found {num_tables}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Tab-separated text lines with correct content (0.5 points)
    # Check that all 4 expected rows appear as paragraphs with tab-separated values.
    # Award partial credit: 0.125 per correct row found.
    try:
        all_para_texts = [p.text for p in doc.paragraphs]
        rows_found = 0
        for expected_row in EXPECTED_ROWS:
            # Normalize: strip whitespace from both sides
            expected_clean = expected_row.strip()
            if any(pt.strip() == expected_clean for pt in all_para_texts):
                rows_found += 1
                print(f"  FOUND: '{expected_clean}'")
            else:
                print(f"  MISSING: '{expected_clean}'")

        if rows_found == 4:
            print(f"PASS: Component 2 — All 4 tab-separated rows found (0.5 pts)")
            total_score += 0.5
        elif rows_found > 0:
            partial = round(rows_found * 0.125, 3)
            print(f"PARTIAL: Component 2 — {rows_found}/4 rows found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No tab-separated rows found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Tab-separated lines correctly positioned AND table removed (0.2 points)
    # In the golden doc, the tab-separated rows appear between the description paragraph
    # and the closing paragraph, with no table present. This compound check ensures
    # both structural correctness AND task completion.
    try:
        para_texts = [p.text.strip() for p in doc.paragraphs]
        # Find indices of the tab-separated rows
        tab_row_indices = []
        for i, pt in enumerate(para_texts):
            if '\t' in pt:
                tab_row_indices.append(i)

        # Check: table removed AND tab rows exist AND they are contiguous
        if num_tables == 0 and len(tab_row_indices) >= 4:
            # Check contiguity: indices should be consecutive
            contiguous = all(
                tab_row_indices[j+1] == tab_row_indices[j] + 1
                for j in range(len(tab_row_indices) - 1)
            )
            if contiguous:
                print(f"PASS: Component 3 — Tab-separated lines are contiguous and table removed (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — Tab-separated lines not contiguous: indices={tab_row_indices}")
        else:
            print(f"FAIL: Component 3 — tables={num_tables}, tab_rows={len(tab_row_indices)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
