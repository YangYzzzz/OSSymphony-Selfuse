"""
Initial Setup: Convert table to tab-separated text
Task ID: writer_tm_011
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_011'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a heading for context
    doc.add_heading('Product Catalog - Table Export', level=1)
    doc.add_paragraph(
        'The following table contains our current product inventory with SKU codes and pricing information.'
    )

    # Create 3x4 table (3 columns, 4 rows including header)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Row 0 (header)
    table.cell(0, 0).text = 'Product'
    table.cell(0, 1).text = 'SKU'
    table.cell(0, 2).text = 'Price'

    # Row 1
    table.cell(1, 0).text = 'Widget A'
    table.cell(1, 1).text = 'W-001'
    table.cell(1, 2).text = '12.99'

    # Row 2
    table.cell(2, 0).text = 'Widget B'
    table.cell(2, 1).text = 'W-002'
    table.cell(2, 2).text = '15.99'

    # Row 3
    table.cell(3, 0).text = 'Widget C'
    table.cell(3, 1).text = 'W-003'
    table.cell(3, 2).text = '9.99'

    # Add a closing paragraph
    doc.add_paragraph(
        'Please update this catalog quarterly to reflect current pricing.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
