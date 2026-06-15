"""
Initial Setup: Create a Writer document with Default Page Style (portrait),
multiple content sections, and an Appendix section at the end.
Task ID: writer_tech_063
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_063'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Ensure default section is portrait A4 with standard margins
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # === Title Page ===
    title = doc.add_heading('Cloud Infrastructure Migration Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by the DevOps Engineering Team')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('March 2026')
    run.font.size = Pt(12)
    run.italic = True

    doc.add_page_break()

    # === Section 1: Executive Summary ===
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_paragraph(
        'This report details the planned migration of our on-premises infrastructure '
        'to a hybrid cloud architecture. The initiative aims to reduce operational costs '
        'by 35% while improving system reliability and scalability. Our analysis covers '
        'the technical requirements, risk assessment, timeline, and resource allocation '
        'needed to execute this transition successfully.'
    )

    doc.add_paragraph(
        'The migration will proceed in three phases over 18 months, beginning with '
        'non-critical development environments and culminating with production workloads. '
        'Total projected investment is $2.4 million, with an expected ROI of 180% over '
        'five years based on current infrastructure spending patterns.'
    )

    doc.add_page_break()

    # === Section 2: Current Architecture ===
    doc.add_heading('2. Current Architecture Overview', level=1)

    doc.add_heading('2.1 Server Infrastructure', level=2)
    doc.add_paragraph(
        'Our current data center houses 247 physical servers across three facilities '
        'in Portland, Austin, and Richmond. The primary facility in Portland handles '
        '68% of all production traffic, with Austin serving as the disaster recovery site '
        'and Richmond hosting development and staging environments.'
    )

    # Table: Server inventory
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Facility', 'Server Count', 'Avg. Utilization', 'Annual Cost']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Portland DC-1', '168', '72%', '$1,245,000'],
        ['Austin DC-2', '52', '34%', '$412,000'],
        ['Richmond DC-3', '27', '58%', '$198,000'],
        ['Total', '247', '55%', '$1,855,000'],
        ['Target (Post-Migration)', '~80', '85%', '$1,200,000'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')  # spacer

    doc.add_heading('2.2 Network Topology', level=2)
    doc.add_paragraph(
        'Inter-facility connectivity is provided by dedicated 10 Gbps MPLS circuits '
        'with automatic failover. Average latency between Portland and Austin is 42ms, '
        'while Portland-to-Richmond averages 28ms. Current bandwidth utilization peaks '
        'at 63% during business hours, leaving adequate headroom for migration traffic.'
    )

    doc.add_page_break()

    # === Section 3: Migration Plan ===
    doc.add_heading('3. Migration Strategy', level=1)

    doc.add_heading('3.1 Phase 1: Development Environments (Months 1-4)', level=2)
    doc.add_paragraph(
        'The first phase targets all development and QA environments currently hosted '
        'in the Richmond facility. This includes 27 servers running a mix of Linux and '
        'Windows workloads. Key activities include:'
    )

    items_phase1 = [
        'Inventory and dependency mapping of all Richmond workloads',
        'Provisioning of cloud VPC with matching network segmentation',
        'Migration of CI/CD pipelines to cloud-native build services',
        'Establishment of hybrid DNS resolution between on-prem and cloud',
        'Performance baseline testing and validation',
    ]
    for item in items_phase1:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2 Phase 2: Staging and DR (Months 5-10)', level=2)
    doc.add_paragraph(
        'Phase 2 focuses on migrating staging environments and establishing cloud-based '
        'disaster recovery. The Austin facility will be decommissioned after successful '
        'validation of cloud DR capabilities. This phase includes database replication '
        'setup, automated failover testing, and compliance verification.'
    )

    doc.add_heading('3.3 Phase 3: Production Workloads (Months 11-18)', level=2)
    doc.add_paragraph(
        'The final phase migrates production services from Portland DC-1. A blue-green '
        'deployment strategy will minimize downtime, with each service cutover planned '
        'during maintenance windows. Critical services including the payment processing '
        'gateway and customer authentication platform will be migrated last with '
        'extended parallel-run periods.'
    )

    doc.add_page_break()

    # === Section 4: Risk Assessment ===
    doc.add_heading('4. Risk Assessment', level=1)

    doc.add_paragraph(
        'A comprehensive risk assessment was conducted across five dimensions: '
        'technical complexity, data integrity, compliance, operational continuity, '
        'and financial exposure. The following table summarizes key risks and '
        'mitigation strategies.'
    )

    risk_table = doc.add_table(rows=5, cols=3)
    risk_table.style = 'Table Grid'
    risk_headers = ['Risk Category', 'Severity', 'Mitigation']
    for i, h in enumerate(risk_headers):
        cell = risk_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    risks = [
        ['Data loss during transfer', 'High', 'Checksum validation + parallel writes'],
        ['Extended downtime', 'Medium', 'Blue-green deployment + rollback plans'],
        ['Compliance gaps', 'Medium', 'Pre-migration audit + continuous monitoring'],
        ['Cost overrun', 'Low', 'Reserved instances + monthly budget reviews'],
    ]
    for r, row_data in enumerate(risks, 1):
        for c, val in enumerate(row_data):
            risk_table.cell(r, c).text = val

    doc.add_page_break()

    # === Appendix Section ===
    doc.add_heading('Appendix', level=1)

    doc.add_heading('A. Detailed Server Inventory', level=2)
    doc.add_paragraph(
        'The following tables provide a complete breakdown of all servers scheduled '
        'for migration, including hardware specifications, operating system versions, '
        'and application dependencies.'
    )

    inv_table = doc.add_table(rows=8, cols=5)
    inv_table.style = 'Table Grid'
    inv_headers = ['Server ID', 'Hostname', 'OS', 'CPU/RAM', 'Primary Application']
    for i, h in enumerate(inv_headers):
        cell = inv_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    inventory = [
        ['SVR-001', 'pdx-web-01', 'RHEL 8.9', '16C/64GB', 'Nginx + Node.js API'],
        ['SVR-002', 'pdx-web-02', 'RHEL 8.9', '16C/64GB', 'Nginx + Node.js API'],
        ['SVR-003', 'pdx-db-01', 'Ubuntu 22.04', '32C/256GB', 'PostgreSQL 15 Primary'],
        ['SVR-004', 'pdx-db-02', 'Ubuntu 22.04', '32C/256GB', 'PostgreSQL 15 Replica'],
        ['SVR-005', 'pdx-cache-01', 'RHEL 8.9', '8C/32GB', 'Redis Cluster Node'],
        ['SVR-006', 'pdx-mq-01', 'Debian 12', '8C/16GB', 'RabbitMQ Broker'],
        ['SVR-007', 'atx-dr-01', 'RHEL 8.9', '16C/64GB', 'DR Orchestrator'],
    ]
    for r, row_data in enumerate(inventory, 1):
        for c, val in enumerate(row_data):
            inv_table.cell(r, c).text = val

    doc.add_paragraph('')

    doc.add_heading('B. Network Configuration Details', level=2)
    doc.add_paragraph(
        'VLAN assignments, firewall rules, and load balancer configurations are '
        'documented below for reference during the migration planning process.'
    )

    doc.add_paragraph('VLAN 100 - Production Web Tier (10.10.100.0/24)', style='List Bullet')
    doc.add_paragraph('VLAN 200 - Production Database Tier (10.10.200.0/24)', style='List Bullet')
    doc.add_paragraph('VLAN 300 - Management Network (10.10.300.0/24)', style='List Bullet')
    doc.add_paragraph('VLAN 400 - DMZ / Public-Facing (10.10.400.0/24)', style='List Bullet')

    doc.add_heading('C. Compliance Checklist', level=2)
    doc.add_paragraph(
        'All migration activities must comply with SOC 2 Type II, HIPAA, and '
        'PCI-DSS requirements. The following checklist must be completed before '
        'each phase go-live:'
    )
    checklist = [
        'Data encryption at rest and in transit verified',
        'Access control lists reviewed and approved by Security team',
        'Penetration testing completed on cloud environment',
        'Incident response plan updated for hybrid architecture',
        'Backup and recovery procedures validated',
        'Audit logging enabled across all cloud services',
    ]
    for item in checklist:
        doc.add_paragraph(item, style='List Number')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
