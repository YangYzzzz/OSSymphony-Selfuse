"""
Reward Script: Create table style for API docs parameter tables
Task ID: writer_tech_067
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Header row background #37474F on all 4 tables
  Component 2 (0.25): Header row white text (FFFFFF) on all 4 tables
  Component 3 (0.25): Alternating data row colors (#FFFFFF / #ECEFF1) on all 4 tables
  Component 4 (0.20): Cell borders 0.5pt #B0BEC5 on all 4 tables
"""

import os
from math import sqrt
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_067'
NUM_TABLES = 4

# Tolerance for color matching (hex string distance)
def hex_color_distance(c1, c2):
    """Compute Euclidean distance between two hex color strings (e.g. '37474F')."""
    try:
        r1, g1, b1 = int(c1[0:2], 16), int(c1[2:4], 16), int(c1[4:6], 16)
        r2, g2, b2 = int(c2[0:2], 16), int(c2[2:4], 16), int(c2[4:6], 16)
        return sqrt((r1-r2)**2 + (g1-g2)**2 + (b1-b2)**2)
    except Exception:
        return 999

COLOR_TOLERANCE = 30  # Allow some tolerance for color matching


def get_cell_shading(cell):
    """Get the fill color from cell shading XML."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        # Also check inside paragraph run properties for shading
        return None
    fill = shd.get(qn('w:fill'))
    return fill


def get_cell_borders(cell):
    """Get border properties from cell. Returns dict of border_name -> (sz, color, val)."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return {}
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        return {}
    borders = {}
    for b in tcBorders:
        tag = b.tag.split('}')[-1]
        sz = b.get(qn('w:sz'))
        color = b.get(qn('w:color'))
        val = b.get(qn('w:val'))
        borders[tag] = (sz, color, val)
    return borders


def get_run_font_color(cell):
    """Get font color from the first non-empty run in the cell."""
    for p in cell.paragraphs:
        for run in p.runs:
            if run.text.strip():
                rgb = run.font.color.rgb
                if rgb is not None:
                    return str(rgb)
                # Also check XML directly
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                rPr = run._element.find('w:rPr', ns)
                if rPr is not None:
                    color_el = rPr.find('w:color', ns)
                    if color_el is not None:
                        return color_el.get(qn('w:val'))
                return None
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    tables = doc.tables
    if len(tables) < NUM_TABLES:
        print("CRITICAL: Expected %d tables, found %d" % (NUM_TABLES, len(tables)))
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Header row background #37474F on all 4 tables (0.30 points)
    try:
        header_bg_pass = 0
        for ti in range(NUM_TABLES):
            table = tables[ti]
            row0 = table.rows[0]
            cells_ok = 0
            for cell in row0.cells:
                fill = get_cell_shading(cell)
                if fill and hex_color_distance(fill.upper(), '37474F') < COLOR_TOLERANCE:
                    cells_ok += 1
            if cells_ok == len(row0.cells):
                header_bg_pass += 1
            else:
                print("FAIL: Component 1 — Table %d header row: %d/%d cells have correct bg" % (ti, cells_ok, len(row0.cells)))

        if header_bg_pass == NUM_TABLES:
            print("PASS: Component 1 — All %d tables have header bg #37474F (0.30 pts)" % NUM_TABLES)
            total_score += 0.30
        elif header_bg_pass > 0:
            partial = 0.30 * (header_bg_pass / NUM_TABLES)
            print("PARTIAL: Component 1 — %d/%d tables have header bg (%.2f pts)" % (header_bg_pass, NUM_TABLES, partial))
            total_score += partial
        else:
            print("FAIL: Component 1 — No tables have header bg #37474F")
    except Exception as e:
        print("ERROR: Component 1 — %s" % e)

    # Component 2: Header row white text (FFFFFF) on all 4 tables (0.25 points)
    try:
        white_text_pass = 0
        for ti in range(NUM_TABLES):
            table = tables[ti]
            row0 = table.rows[0]
            cells_ok = 0
            total_text_cells = 0
            for cell in row0.cells:
                color = get_run_font_color(cell)
                if color is not None:
                    total_text_cells += 1
                    if hex_color_distance(color.upper(), 'FFFFFF') < COLOR_TOLERANCE:
                        cells_ok += 1
            if total_text_cells > 0 and cells_ok == total_text_cells:
                white_text_pass += 1
            else:
                print("FAIL: Component 2 — Table %d header: %d/%d text cells have white text" % (ti, cells_ok, total_text_cells))

        if white_text_pass == NUM_TABLES:
            print("PASS: Component 2 — All %d tables have white header text (0.25 pts)" % NUM_TABLES)
            total_score += 0.25
        elif white_text_pass > 0:
            partial = 0.25 * (white_text_pass / NUM_TABLES)
            print("PARTIAL: Component 2 — %d/%d tables have white header text (%.2f pts)" % (white_text_pass, NUM_TABLES, partial))
            total_score += partial
        else:
            print("FAIL: Component 2 — No tables have white header text")
    except Exception as e:
        print("ERROR: Component 2 — %s" % e)

    # Component 3: Alternating data row colors (#FFFFFF / #ECEFF1) on all 4 tables (0.25 points)
    try:
        alt_color_pass = 0
        for ti in range(NUM_TABLES):
            table = tables[ti]
            num_rows = len(table.rows)
            if num_rows < 2:
                continue
            rows_ok = 0
            data_rows = num_rows - 1  # exclude header
            for ri in range(1, num_rows):
                # Odd data rows (ri=1,3,5...) should be #FFFFFF
                # Even data rows (ri=2,4,6...) should be #ECEFF1
                expected = 'FFFFFF' if ri % 2 == 1 else 'ECEFF1'
                row = table.rows[ri]
                cell_fill = get_cell_shading(row.cells[0])
                if cell_fill and hex_color_distance(cell_fill.upper(), expected) < COLOR_TOLERANCE:
                    rows_ok += 1
                else:
                    print("FAIL: Component 3 — Table %d row %d: expected %s, got %s" % (ti, ri, expected, cell_fill))
            if rows_ok == data_rows:
                alt_color_pass += 1

        if alt_color_pass == NUM_TABLES:
            print("PASS: Component 3 — All %d tables have alternating row colors (0.25 pts)" % NUM_TABLES)
            total_score += 0.25
        elif alt_color_pass > 0:
            partial = 0.25 * (alt_color_pass / NUM_TABLES)
            print("PARTIAL: Component 3 — %d/%d tables have alternating colors (%.2f pts)" % (alt_color_pass, NUM_TABLES, partial))
            total_score += partial
        else:
            print("FAIL: Component 3 — No tables have alternating row colors")
    except Exception as e:
        print("ERROR: Component 3 — %s" % e)

    # Component 4: Cell borders 0.5pt #B0BEC5 on all 4 tables (0.20 points)
    # 0.5pt = sz=4 in OOXML (eighths of a point)
    try:
        border_pass = 0
        for ti in range(NUM_TABLES):
            table = tables[ti]
            bad_cells = 0
            total_cells = 0
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    total_cells += 1
                    borders = get_cell_borders(cell)
                    if not borders:
                        bad_cells += 1
                        if ri == 0 and ci == 0:
                            print("FAIL: Component 4 — Table %d [%d,%d]: no cell borders" % (ti, ri, ci))
                        continue
                    # Check at least top/left/bottom/right exist with correct color
                    cell_valid = (len({'top', 'left', 'bottom', 'right'} - set(borders.keys())) == 0)
                    if cell_valid:
                        for side in ['top', 'left', 'bottom', 'right']:
                            sz, color, val = borders[side]
                            if sz is None or int(sz) != 4:
                                cell_valid = False
                                break
                            if color is None or hex_color_distance(color.upper(), 'B0BEC5') > COLOR_TOLERANCE:
                                cell_valid = False
                                break
                    if not cell_valid:
                        bad_cells += 1
            if total_cells > 0 and bad_cells == 0:
                border_pass += 1

        if border_pass == NUM_TABLES:
            print("PASS: Component 4 — All %d tables have 0.5pt #B0BEC5 borders (0.20 pts)" % NUM_TABLES)
            total_score += 0.20
        elif border_pass > 0:
            partial = 0.20 * (border_pass / NUM_TABLES)
            print("PARTIAL: Component 4 — %d/%d tables have correct borders (%.2f pts)" % (border_pass, NUM_TABLES, partial))
            total_score += partial
        else:
            print("FAIL: Component 4 — No tables have correct 0.5pt #B0BEC5 borders")
    except Exception as e:
        print("ERROR: Component 4 — %s" % e)

    final_score = min(total_score, 1.0)
    print("\nScore: %.2f/1.0" % total_score)
    print("REWARD: %.1f" % final_score)
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print("PERSIST_WARN: save hook failed: %s" % e)


# Entry point
file_path = os.path.join(WORKDIR, '%s.docx' % TASK_ID)
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
