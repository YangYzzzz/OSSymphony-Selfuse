"""
Reward Script: Employee Performance Review Form
Task ID: writer_wf_036
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15): Title "Annual Performance Review"
  Component 2 (0.20): Employee Information table with 5 fields
  Component 3 (0.25): Performance Ratings table with 6 criteria + Rating/Comments columns
  Component 4 (0.20): Four additional sections (Goals Achieved, Areas for Improvement, Development Plan, Overall Rating)
  Component 5 (0.10): Signatures section with Employee and Manager areas
  Component 6 (0.10): Document structure — proper headings usage
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_036'


def persist_app_state(domain: str):
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_text = ' '.join(p.text.strip().lower() for p in doc.paragraphs if p.text.strip())

    # Component 1: Title "Annual Performance Review" (0.15 points)
    try:
        has_title = False
        for p in doc.paragraphs:
            if 'annual performance review' in p.text.strip().lower():
                has_title = True
                break
        if has_title:
            print(f"PASS: Component 1 — Title 'Annual Performance Review' found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Title 'Annual Performance Review' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Employee Information table with 5 fields (0.20 points)
    # Required fields: Name, ID, Department, Manager, Review Period
    try:
        required_info_fields = ['name', 'id', 'department', 'manager', 'review period']
        found_fields = 0
        for table in doc.tables:
            table_text = ' '.join(cell.text.strip().lower() for row in table.rows for cell in row.cells)
            field_matches = sum(1 for f in required_info_fields if f in table_text)
            found_fields = max(found_fields, field_matches)

        if found_fields >= 4:
            print(f"PASS: Component 2 — Employee info table found with {found_fields}/5 fields (0.20 pts)")
            total_score += 0.20
        elif found_fields >= 3:
            print(f"PARTIAL: Component 2 — Employee info table found with {found_fields}/5 fields (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — Employee info table insufficient: {found_fields}/5 fields")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Performance Ratings table with 6 criteria + Rating/Comments (0.25 points)
    # Criteria: Job Knowledge, Quality of Work, Productivity, Communication, Teamwork, Initiative
    try:
        required_criteria = ['job knowledge', 'quality of work', 'productivity',
                             'communication', 'teamwork', 'initiative']
        best_criteria_count = 0
        has_rating_col = False
        has_comments_col = False

        for table in doc.tables:
            table_text_full = ' '.join(cell.text.strip().lower() for row in table.rows for cell in row.cells)
            criteria_found = sum(1 for c in required_criteria if c in table_text_full)
            if criteria_found > best_criteria_count:
                best_criteria_count = criteria_found
                # Check for Rating and Comments columns
                if len(table.rows) > 0:
                    header_text = ' '.join(cell.text.strip().lower() for cell in table.rows[0].cells)
                    has_rating_col = 'rating' in header_text
                    has_comments_col = 'comment' in header_text

        score_3 = 0.0
        # Up to 0.15 for criteria presence
        if best_criteria_count >= 5:
            score_3 += 0.15
        elif best_criteria_count >= 3:
            score_3 += 0.08

        # Up to 0.05 for Rating column
        if has_rating_col:
            score_3 += 0.05

        # Up to 0.05 for Comments column
        if has_comments_col:
            score_3 += 0.05

        if score_3 > 0:
            print(f"PASS: Component 3 — Performance Ratings: {best_criteria_count}/6 criteria, "
                  f"rating_col={has_rating_col}, comments_col={has_comments_col} ({score_3} pts)")
            total_score += score_3
        else:
            print(f"FAIL: Component 3 — Performance Ratings table not found or missing criteria")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Four additional sections (0.20 points)
    # Goals Achieved, Areas for Improvement, Development Plan, Overall Rating
    try:
        required_sections = ['goals achieved', 'areas for improvement', 'development plan', 'overall rating']
        found_sections = 0
        for section_name in required_sections:
            for p in doc.paragraphs:
                if section_name in p.text.strip().lower():
                    found_sections += 1
                    break

        if found_sections >= 4:
            print(f"PASS: Component 4 — All 4 sections found ({found_sections}/4) (0.20 pts)")
            total_score += 0.20
        elif found_sections >= 2:
            partial = round(0.05 * found_sections, 2)
            print(f"PARTIAL: Component 4 — {found_sections}/4 sections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {found_sections}/4 sections found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Signatures section with Employee and Manager (0.10 points)
    try:
        has_signature_section = False
        has_employee_sig = False
        has_manager_sig = False

        # Check paragraphs for "signatures" heading
        for p in doc.paragraphs:
            if 'signature' in p.text.strip().lower():
                has_signature_section = True
                break

        # Check tables for dual signature areas
        for table in doc.tables:
            table_text = ' '.join(cell.text.strip().lower() for row in table.rows for cell in row.cells)
            if 'employee' in table_text and ('manager' in table_text or 'supervisor' in table_text):
                if 'signature' in table_text or 'sign' in table_text or 'date' in table_text:
                    has_employee_sig = True
                    has_manager_sig = True

        # Also check paragraph text for signature patterns without a table
        if not (has_employee_sig and has_manager_sig):
            sig_text = ' '.join(p.text.strip().lower() for p in doc.paragraphs)
            if ('employee' in sig_text and 'signature' in sig_text and
                ('manager' in sig_text or 'supervisor' in sig_text)):
                has_employee_sig = True
                has_manager_sig = True

        if has_signature_section and has_employee_sig and has_manager_sig:
            print(f"PASS: Component 5 — Signatures section with employee and manager areas (0.10 pts)")
            total_score += 0.10
        elif has_signature_section:
            print(f"PARTIAL: Component 5 — Signatures heading found but missing dual areas (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 — No signatures section found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Document structure — headings used (0.10 points)
    # The form should use headings (not just plain text) for section organization
    try:
        heading_count = 0
        for p in doc.paragraphs:
            if p.style and 'Heading' in p.style.name:
                heading_count += 1

        if heading_count >= 4:
            print(f"PASS: Component 6 — {heading_count} headings used for structure (0.10 pts)")
            total_score += 0.10
        elif heading_count >= 2:
            print(f"PARTIAL: Component 6 — {heading_count} headings used (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — Only {heading_count} headings found (expected >= 4)")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
