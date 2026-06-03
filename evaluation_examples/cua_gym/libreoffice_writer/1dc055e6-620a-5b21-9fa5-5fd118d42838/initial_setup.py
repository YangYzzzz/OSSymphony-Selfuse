"""
Initial Setup: Create a Writer document with cover page text at top with default spacing.
Task ID: writer_rd_069
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_069'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
LOGO_PATH = f'{WORKDIR}/company_logo.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_logo():
    """Create a simple company logo image."""
    img = Image.new('RGB', (400, 400), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    # Draw a blue circle as logo base
    draw.ellipse([50, 50, 350, 350], fill=(31, 78, 121), outline=(21, 52, 80), width=4)
    # Draw letters "AF" in center
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 120)
    except (OSError, IOError):
        font = ImageFont.load_default()
    draw.text((110, 120), "AF", fill=(255, 255, 255), font=font)
    img.save(LOGO_PATH)
    print(f'Logo created: {LOGO_PATH}')


def create_initial():
    doc = Document()

    # Plain text lines stacked at top with default formatting (no centering, no size changes)
    lines = [
        'Annual Financial Report',
        'Fiscal Year 2024-2025',
        'Prepared by: Finance Department',
        'January 2025',
    ]

    for line_text in lines:
        para = doc.add_paragraph()
        run = para.add_run(line_text)
        # Default formatting: no bold, no centering, default size (~11pt)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_logo()
create_initial()
