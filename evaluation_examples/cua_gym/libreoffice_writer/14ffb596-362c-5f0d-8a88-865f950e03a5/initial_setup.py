"""
Initial Setup: Magazine-style article for Drop Cap task
Task ID: writer_rd_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font to Liberation Serif 11pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Liberation Serif'
    font.size = Pt(11)

    # Add a title
    heading = doc.add_heading('The Renaissance of Urban Gardens', level=1)
    for run in heading.runs:
        run.font.name = 'Liberation Serif'

    # Paragraph 1 - starts with 'T'
    p1_text = (
        "Throughout the last decade, urban gardening has transformed from a niche hobby "
        "into a widespread movement reshaping cities across the globe. From rooftop farms "
        "in Brooklyn to vertical gardens in Singapore, communities are discovering that "
        "growing food locally is not only sustainable but also deeply rewarding. Municipal "
        "planners have taken note, allocating previously neglected parcels of land for "
        "community garden projects that bring neighbors together and provide fresh produce "
        "to food-insecure neighborhoods."
    )
    para1 = doc.add_paragraph()
    run1 = para1.add_run(p1_text)
    run1.font.name = 'Liberation Serif'
    run1.font.size = Pt(11)

    # Paragraph 2
    p2_text = (
        "In cities like Detroit, where vacant lots once symbolized economic decline, "
        "community gardens have become beacons of renewal. Organizations such as the "
        "Michigan Urban Farming Initiative have converted abandoned properties into "
        "thriving agricultural spaces that produce thousands of pounds of organic "
        "vegetables each season. These gardens serve as outdoor classrooms where children "
        "learn about nutrition, biology, and environmental stewardship firsthand."
    )
    para2 = doc.add_paragraph()
    run2 = para2.add_run(p2_text)
    run2.font.name = 'Liberation Serif'
    run2.font.size = Pt(11)

    # Paragraph 3
    p3_text = (
        "The economic benefits extend far beyond the harvest itself. Property values "
        "in neighborhoods with well-maintained community gardens have risen by an average "
        "of twelve percent according to a 2024 study by the National Association of "
        "Realtors. Local businesses report increased foot traffic, and farmers markets "
        "connected to urban gardens generate an estimated three billion dollars annually "
        "in direct sales across the United States alone."
    )
    para3 = doc.add_paragraph()
    run3 = para3.add_run(p3_text)
    run3.font.name = 'Liberation Serif'
    run3.font.size = Pt(11)

    # Paragraph 4
    p4_text = (
        "Technology is accelerating this green revolution. Smart irrigation systems "
        "powered by artificial intelligence can reduce water usage by up to forty percent "
        "while maintaining optimal soil moisture levels. Hydroponic towers in repurposed "
        "shipping containers allow year-round cultivation regardless of climate, and "
        "solar-powered LED grow lights have made indoor farming viable even in the "
        "northernmost cities where winter daylight is scarce."
    )
    para4 = doc.add_paragraph()
    run4 = para4.add_run(p4_text)
    run4.font.name = 'Liberation Serif'
    run4.font.size = Pt(11)

    # Paragraph 5
    p5_text = (
        "As we look ahead to the next decade, experts predict that urban agriculture "
        "will become an integral part of city infrastructure planning. Architects are "
        "already designing buildings with integrated growing spaces, and policymakers "
        "are drafting legislation to incentivize green rooftops and edible landscaping. "
        "The vision of a city where fresh herbs grow on every balcony and fruit trees "
        "line every boulevard may soon become reality."
    )
    para5 = doc.add_paragraph()
    run5 = para5.add_run(p5_text)
    run5.font.name = 'Liberation Serif'
    run5.font.size = Pt(11)

    # Set paragraph formatting - justified text for magazine style
    for para in doc.paragraphs[1:]:  # skip heading
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para.paragraph_format.space_after = Pt(8)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
