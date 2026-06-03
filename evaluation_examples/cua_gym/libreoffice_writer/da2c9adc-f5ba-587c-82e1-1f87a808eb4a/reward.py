"""
Reward Script: Format superscripts in 'a2 + b2 = c2' and subscript in 'log10(x)'
Task ID: writer_txtfmt_044
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): In paragraph containing 'a2 + b2 = c2', each '2' run is superscript=True (3 runs)
  Component 2 (0.5 pts): In paragraph containing 'log10(x)', the '10' run is subscript=True
Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_txtfmt_044'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Superscript '2's in 'a2 + b2 = c2' paragraph (0.5 points)
    # The Pythagorean theorem paragraph (Para 0) must have each '2' as superscript=True.
    # In initial_env, all '2' runs have superscript=None (not superscript).
    # In golden_env, all three '2' runs have superscript=True.
    try:
        pythagorean_para = None
        for para in doc.paragraphs:
            if 'a2 + b2 = c2' in para.text or ('a2' in para.text and 'b2' in para.text and 'c2' in para.text):
                pythagorean_para = para
                break

        if pythagorean_para is None:
            print("FAIL: Component 1 — Could not find paragraph containing 'a2 + b2 = c2'")
        else:
            # Find runs that contain only '2' and check for superscript
            superscript_runs = []
            non_superscript_twos = []
            for run in pythagorean_para.runs:
                if run.text == '2':
                    if run.font.superscript is True:
                        superscript_runs.append(run.text)
                    else:
                        non_superscript_twos.append(f"text={repr(run.text)} superscript={run.font.superscript}")

            if len(superscript_runs) == 3 and len(non_superscript_twos) == 0:
                print(f"PASS: Component 1 — All 3 '2' runs in Pythagorean paragraph are superscript=True (0.5 pts)")
                total_score += 0.5
            else:
                print(f"FAIL: Component 1 — Expected 3 superscript '2' runs, found {len(superscript_runs)} superscript; "
                      f"non-superscript '2' runs: {non_superscript_twos}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Subscript '10' in 'log10(x)' paragraph (0.5 points)
    # The logarithm paragraph (Para 2) must have '10' run with subscript=True.
    # In initial_env, '10' run has subscript=None (not subscript).
    # In golden_env, '10' run has subscript=True.
    try:
        log_para = None
        for para in doc.paragraphs:
            if 'log10(x)' in para.text or ('log' in para.text and '10' in para.text and '(x)' in para.text):
                log_para = para
                break

        if log_para is None:
            print("FAIL: Component 2 — Could not find paragraph containing 'log10(x)'")
        else:
            # Find '10' runs and check for subscript
            subscript_ten_count = 0
            non_subscript_tens = []
            for run in log_para.runs:
                if run.text == '10':
                    if run.font.subscript is True:
                        subscript_ten_count += 1
                    else:
                        non_subscript_tens.append(f"text={repr(run.text)} subscript={run.font.subscript}")

            if subscript_ten_count > 0:
                print(f"PASS: Component 2 — '10' run in logarithm paragraph is subscript=True (0.5 pts)")
                total_score += 0.5
            else:
                print(f"FAIL: Component 2 — Expected '10' run to be subscript=True; "
                      f"non-subscript '10' runs: {non_subscript_tens}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM Desktop
file_path = f'{WORKDIR}/math_homework.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
