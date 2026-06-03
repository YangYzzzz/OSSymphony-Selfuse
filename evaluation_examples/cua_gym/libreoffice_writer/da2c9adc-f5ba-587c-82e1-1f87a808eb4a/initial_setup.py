"""
Initial Setup: Math homework document with superscript/subscript task
Task ID: writer_txtfmt_044
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'  # VM path — task specifies ~/Desktop
TASK_ID = 'math_homework'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_run_font(run, font_name='Liberation Sans', size_pt=12):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    # Ensure no super/subscript
    run.font.superscript = False
    run.font.subscript = False


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Paragraph 1: Pythagorean theorem with plain '2' (NOT superscript)
    para1 = doc.add_paragraph()

    run1a = para1.add_run('The Pythagorean theorem states that a')
    set_run_font(run1a)
    run1b = para1.add_run('2')
    set_run_font(run1b)  # plain '2', NOT superscript
    run1c = para1.add_run(' + b')
    set_run_font(run1c)
    run1d = para1.add_run('2')
    set_run_font(run1d)  # plain '2', NOT superscript
    run1e = para1.add_run(' = c')
    set_run_font(run1e)
    run1f = para1.add_run('2')
    set_run_font(run1f)  # plain '2', NOT superscript
    run1g = para1.add_run(' for any right triangle.')
    set_run_font(run1g)

    # Paragraph 2: Explanatory content
    para2 = doc.add_paragraph()
    run2 = para2.add_run(
        'This fundamental relationship is one of the most well-known results in Euclidean geometry, '
        'named after the ancient Greek mathematician Pythagoras of Samos.'
    )
    set_run_font(run2)

    # Paragraph 3: Logarithm with plain '10' (NOT subscript)
    para3 = doc.add_paragraph()
    run3a = para3.add_run('The common logarithm log')
    set_run_font(run3a)
    run3b = para3.add_run('10')
    set_run_font(run3b)  # plain '10', NOT subscript
    run3c = para3.add_run('(x) is defined as the power to which 10 must be raised.')
    set_run_font(run3c)

    # Paragraph 4: Additional content for realism
    para4 = doc.add_paragraph()
    run4 = para4.add_run(
        'For example, log10(100) = 2 because 10 raised to the power 2 equals 100. '
        'Logarithms are widely used in science and engineering calculations.'
    )
    set_run_font(run4)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
