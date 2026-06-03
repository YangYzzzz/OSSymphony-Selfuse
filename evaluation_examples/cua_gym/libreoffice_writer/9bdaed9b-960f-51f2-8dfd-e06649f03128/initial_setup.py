"""
Initial Setup: Customer Satisfaction Survey with 8 questions (no form controls)
Task ID: writer_rd_038
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_038'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Customer Satisfaction Survey', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Introduction ---
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        'Thank you for taking the time to complete our customer satisfaction survey. '
        'Your feedback helps us improve our products and services. Please answer each '
        'question honestly by selecting Yes or No.'
    )
    intro_run.font.size = Pt(11)
    intro_run.font.name = 'Calibri'
    intro.paragraph_format.space_after = Pt(12)

    # --- Section Header ---
    section_header = doc.add_heading('Survey Questions', level=1)
    section_header.paragraph_format.space_before = Pt(6)
    section_header.paragraph_format.space_after = Pt(12)

    # --- 8 Survey Questions ---
    questions = [
        '1. Are you satisfied with our overall service quality?',
        '2. Would you recommend our company to a friend or colleague?',
        '3. Did our staff respond to your inquiries in a timely manner?',
        '4. Were you satisfied with the range of products available?',
        '5. Do you feel our pricing is fair and competitive?',
        '6. Was the checkout process smooth and easy to complete?',
        '7. Have you experienced any issues with product delivery?',
        '8. Would you consider purchasing from us again in the future?',
    ]

    for q_text in questions:
        para = doc.add_paragraph()
        run = para.add_run(q_text)
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)

    # --- Closing ---
    doc.add_paragraph()  # spacer
    closing = doc.add_paragraph()
    closing_run = closing.add_run(
        'Thank you for your participation. Your responses are confidential and will '
        'be used solely to improve our services.'
    )
    closing_run.font.size = Pt(10)
    closing_run.font.italic = True
    closing_run.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
