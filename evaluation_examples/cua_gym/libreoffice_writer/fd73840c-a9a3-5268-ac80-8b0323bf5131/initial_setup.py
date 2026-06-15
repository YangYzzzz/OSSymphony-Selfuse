"""
Initial Setup: Table of contents draft with section names and page numbers, no tabstops
Task ID: osworld_writer_tabstop_002
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_tabstop_002'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title paragraph
    title = doc.add_paragraph()
    title_run = title.add_run("Table of Contents")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Empty paragraph after title
    doc.add_paragraph()

    # 12 table-of-contents entries: section name followed by space and page number
    # No tab characters, no tabstops — plain left-aligned text
    toc_entries = [
        ("Introduction", 1),
        ("Background", 3),
        ("Literature Review", 5),
        ("Methodology", 7),
        ("Data Collection", 9),
        ("Analysis Framework", 11),
        ("Results", 12),
        ("Discussion", 15),
        ("Implications", 18),
        ("Limitations", 20),
        ("Conclusion", 22),
        ("References", 24),
    ]

    for section_name, page_num in toc_entries:
        para = doc.add_paragraph()
        # Plain text: name and page number separated by a space — no tab
        run = para.add_run(f"{section_name} {page_num}")
        run.font.size = Pt(12)
        # No tabstops set — paragraph_format.tab_stops is empty

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
