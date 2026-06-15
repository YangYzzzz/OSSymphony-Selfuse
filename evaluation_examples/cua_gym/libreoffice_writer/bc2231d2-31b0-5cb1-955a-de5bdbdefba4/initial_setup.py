"""
Initial Setup: Cover letter with 4 paragraphs for paragraph deletion task
Task ID: writer_edit_053
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'cover_letter'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Paragraph 1: Opening paragraph
    p1 = doc.add_paragraph(
        "Dear Hiring Manager, I am writing to express my strong interest in the "
        "Software Engineer position at TechVision Solutions. With a passion for "
        "building scalable systems and a proven track record of delivering quality "
        "software, I am confident I would be a valuable addition to your team."
    )

    # Paragraph 2: The paragraph to be deleted (exact text from task context)
    p2 = doc.add_paragraph(
        "I recently graduated from Springfield University with a degree in Computer "
        "Science, where I maintained a 3.8 GPA and participated in several hackathons."
    )

    # Paragraph 3: Skills/experience paragraph
    p3 = doc.add_paragraph(
        "During my academic career and internships, I gained hands-on experience with "
        "Python, Java, and cloud infrastructure on AWS. I led a team of four in developing "
        "a real-time data pipeline that reduced processing latency by 40%, demonstrating "
        "my ability to deliver impactful results under tight deadlines."
    )

    # Paragraph 4: Closing paragraph
    p4 = doc.add_paragraph(
        "I would welcome the opportunity to discuss how my background aligns with "
        "TechVision Solutions' goals. Thank you for considering my application. "
        "I look forward to speaking with you at your earliest convenience. "
        "Sincerely, Alex Rivera."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')
    print(f'Paragraphs: {len([p for p in doc.paragraphs if p.text.strip()])} non-empty paragraphs')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
