"""
Initial Setup: Set up table numbering to follow chapter-based pattern
Task ID: writer_acad_056
Domain: libreoffice_writer

Creates a Writer document with 3 chapters (Heading 1 with numbering),
each containing 2 tables with simple sequential captions (Table 1 through Table 6).
The task is for the agent to change these to chapter-based numbering (Table 1.1, etc.).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_056'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_chapter_heading(doc, chapter_num, title):
    """Add a chapter heading with explicit numbering prefix."""
    heading = doc.add_heading(level=1)
    heading.clear()
    run = heading.add_run(f"Chapter {chapter_num}: {title}")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    return heading


def add_body_text(doc, text):
    """Add a body paragraph with realistic academic content."""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)
    for run in para.runs:
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
    return para


def add_table_with_caption(doc, caption_text, headers, data):
    """Add a table with a caption paragraph above it."""
    # Caption paragraph
    caption_para = doc.add_paragraph()
    caption_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption_para.paragraph_format.space_before = Pt(12)
    caption_para.paragraph_format.space_after = Pt(6)
    caption_run = caption_para.add_run(caption_text)
    caption_run.bold = True
    caption_run.font.size = Pt(10)
    caption_run.font.name = "Times New Roman"

    # Create the table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

    # Data rows
    for row_data in data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"

    # Spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)

    return table


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # --- Document Title ---
    title = doc.add_heading("A Comparative Analysis of Machine Learning Approaches for Urban Traffic Flow Prediction", level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(18)

    # Abstract
    abstract_heading = doc.add_heading("Abstract", level=2)
    add_body_text(doc,
        "This study presents a comprehensive evaluation of machine learning techniques "
        "applied to urban traffic flow prediction. We analyze data collected from 47 "
        "sensor stations across three metropolitan areas over a 24-month period. "
        "Our findings demonstrate that hybrid ensemble methods outperform single-model "
        "approaches by 12-18% in terms of mean absolute percentage error (MAPE)."
    )

    # ============================================================
    # CHAPTER 1: Introduction and Literature Review
    # ============================================================
    add_chapter_heading(doc, 1, "Introduction and Literature Review")

    add_body_text(doc,
        "Traffic flow prediction remains a critical challenge in intelligent transportation "
        "systems (ITS). Accurate short-term predictions enable dynamic signal control, "
        "route guidance, and congestion management. Traditional statistical methods such as "
        "ARIMA and its variants have been widely used, but their linear assumptions limit "
        "performance under non-stationary conditions (Zhang et al., 2023)."
    )

    add_body_text(doc,
        "Recent advances in deep learning have opened new avenues for capturing complex "
        "spatial-temporal dependencies in traffic data. Convolutional neural networks (CNNs), "
        "recurrent neural networks (RNNs), and graph neural networks (GNNs) have all shown "
        "promising results in various traffic prediction tasks."
    )

    # Table 1: Literature comparison
    add_table_with_caption(doc,
        "Table 1: Summary of Recent Traffic Prediction Studies",
        ["Study", "Method", "Dataset", "MAPE (%)"],
        [
            ["Zhang et al. (2023)", "LSTM-Attention", "PeMS-BAY", "4.32"],
            ["Liu & Wang (2022)", "ST-GCN", "METR-LA", "3.87"],
            ["Park et al. (2023)", "Transformer", "Seoul Metro", "5.14"],
            ["Chen & Li (2024)", "CNN-BiLSTM", "Beijing Ring", "3.95"],
            ["Gupta et al. (2023)", "XGBoost Ensemble", "Mumbai Traffic", "6.21"],
            ["Kim & Tanaka (2024)", "GAN-LSTM", "Tokyo Highway", "4.08"],
        ]
    )

    add_body_text(doc,
        "As shown in Table 1, graph-based methods generally achieve lower error rates "
        "on spatially-structured datasets. However, the computational overhead of these "
        "methods remains a concern for real-time deployment scenarios."
    )

    # Table 2: Sensor deployment summary
    add_table_with_caption(doc,
        "Table 2: Sensor Deployment Summary Across Study Areas",
        ["Metropolitan Area", "Sensors", "Coverage (km)", "Data Period"],
        [
            ["Metro Area A (Riverside)", "18", "42.5", "Jan 2022 - Dec 2023"],
            ["Metro Area B (Lakewood)", "15", "38.7", "Mar 2022 - Feb 2024"],
            ["Metro Area C (Hilltop)", "14", "35.2", "Jun 2022 - May 2024"],
        ]
    )

    # ============================================================
    # CHAPTER 2: Methodology
    # ============================================================
    add_chapter_heading(doc, 2, "Methodology")

    add_body_text(doc,
        "Our methodology combines three complementary approaches: (a) a temporal attention "
        "mechanism for capturing periodic patterns, (b) a spatial graph convolution layer "
        "for modeling inter-sensor relationships, and (c) an ensemble aggregation strategy "
        "that dynamically weights model outputs based on recent prediction accuracy."
    )

    add_body_text(doc,
        "The data preprocessing pipeline includes outlier detection using the modified "
        "Z-score method (threshold = 3.5), missing value imputation via spatially-weighted "
        "k-nearest neighbor interpolation, and temporal normalization to account for "
        "seasonal and weekly cyclical patterns."
    )

    # Table 3: Hyperparameter configurations
    add_table_with_caption(doc,
        "Table 3: Hyperparameter Configurations for Each Model",
        ["Parameter", "LSTM-Att", "ST-GCN", "Transformer", "Ensemble"],
        [
            ["Learning Rate", "0.001", "0.0005", "0.0001", "N/A"],
            ["Batch Size", "64", "32", "128", "N/A"],
            ["Hidden Units", "256", "128", "512", "N/A"],
            ["Dropout Rate", "0.3", "0.2", "0.1", "N/A"],
            ["Epochs", "200", "150", "100", "N/A"],
            ["Weight Decay", "1e-5", "1e-4", "1e-5", "N/A"],
            ["Attention Heads", "8", "4", "16", "N/A"],
        ]
    )

    add_body_text(doc,
        "All experiments were conducted on a computing cluster with 4x NVIDIA A100 GPUs "
        "(80GB VRAM each). Training times ranged from 2.5 hours for the LSTM-Attention "
        "model to 8.3 hours for the full ensemble pipeline."
    )

    # Table 4: Feature engineering summary
    add_table_with_caption(doc,
        "Table 4: Feature Engineering Summary",
        ["Feature Category", "Features", "Dimensionality"],
        [
            ["Temporal", "Hour, day-of-week, month, holiday flag", "7"],
            ["Spatial", "Adjacency matrix, distance, road type", "47 x 47"],
            ["Traffic", "Flow, speed, occupancy (5-min intervals)", "3 x 288"],
            ["Weather", "Temperature, precipitation, visibility", "3"],
            ["Events", "Incidents, construction, special events", "3"],
        ]
    )

    # ============================================================
    # CHAPTER 3: Results and Discussion
    # ============================================================
    add_chapter_heading(doc, 3, "Results and Discussion")

    add_body_text(doc,
        "The experimental results are summarized across multiple prediction horizons "
        "(15-min, 30-min, and 60-min ahead). We evaluate each model using MAPE, RMSE, "
        "and MAE metrics on a held-out test set comprising the final 3 months of data "
        "from each metropolitan area."
    )

    # Table 5: Main results
    add_table_with_caption(doc,
        "Table 5: Prediction Performance Comparison (15-min Horizon)",
        ["Model", "MAPE (%)", "RMSE", "MAE", "Training Time (h)"],
        [
            ["ARIMA", "8.74", "15.32", "10.21", "0.3"],
            ["LSTM-Attention", "4.28", "8.45", "5.67", "2.5"],
            ["ST-GCN", "3.91", "7.82", "5.12", "4.1"],
            ["Transformer", "4.05", "8.11", "5.34", "3.7"],
            ["XGBoost", "5.63", "10.74", "7.18", "0.8"],
            ["Hybrid Ensemble", "3.42", "6.95", "4.56", "8.3"],
        ]
    )

    add_body_text(doc,
        "The hybrid ensemble achieves the best performance across all metrics, with a "
        "MAPE of 3.42% for the 15-minute prediction horizon. This represents an improvement "
        "of 12.5% over the best individual model (ST-GCN). The performance gap widens "
        "at longer prediction horizons, reaching 18.2% improvement at the 60-minute mark."
    )

    # Table 6: Statistical significance
    add_table_with_caption(doc,
        "Table 6: Statistical Significance Tests (Paired t-test, p-values)",
        ["Comparison", "15-min", "30-min", "60-min"],
        [
            ["Ensemble vs ARIMA", "< 0.001", "< 0.001", "< 0.001"],
            ["Ensemble vs LSTM-Att", "0.003", "0.001", "< 0.001"],
            ["Ensemble vs ST-GCN", "0.021", "0.008", "0.002"],
            ["Ensemble vs Transformer", "0.012", "0.005", "0.001"],
            ["Ensemble vs XGBoost", "< 0.001", "< 0.001", "< 0.001"],
        ]
    )

    add_body_text(doc,
        "All pairwise comparisons between the hybrid ensemble and individual models "
        "are statistically significant at the 0.05 level (Table 6). The Diebold-Mariano "
        "test confirms these findings, indicating that the forecast accuracy improvements "
        "are not attributable to random variation."
    )

    # Save document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
