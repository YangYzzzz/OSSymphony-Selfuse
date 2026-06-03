"""
Initial Setup: Document with common errors for AutoCorrect application
Task ID: writer_frd_054
Domain: libreoffice_writer

Errors planted (15 total):
  - 4x '(c)' that should become copyright symbol
  - 4x 'teh' that should become 'the'
  - 4x lowercase sentence starts (after period+space)
  - 3x straight quote pairs that should become smart quotes
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_054'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Liberation Serif'
    font.size = Pt(12)

    # Title
    doc.add_heading('Quarterly Marketing Report', level=1)

    # Error legend (for counting):
    # (c) = errors #1-4
    # teh = errors #5-8
    # lowercase sentence start = errors #9-12
    # straight quotes = errors #13-15

    # --- Paragraph 1 ---
    # Errors: teh (#5), (c) (#1)
    doc.add_paragraph(
        'teh marketing department has completed its Q3 review. '
        'Our flagship product, branded under the company trademark (c) 2024, '
        'continued to outperform expectations across all regions.'
    )

    # --- Paragraph 2 ---
    # Errors: lowercase start "according" (#9), straight quotes (#13), teh (#6)
    doc.add_paragraph(
        'according to the latest survey, customer satisfaction reached 94%. '
        '"We are thrilled with these results," said Director Elena Vasquez. '
        'teh team worked incredibly hard this quarter to achieve these numbers.'
    )

    # --- Paragraph 3 ---
    # Errors: (c) (#2), lowercase start "we" (#10)
    doc.add_paragraph(
        'The intellectual property portfolio expanded significantly. '
        'we filed three new trademark applications, each bearing the (c) '
        'symbol for brand protection across North America and Europe.'
    )

    # Subheading
    doc.add_heading('Budget Summary', level=2)

    # --- Paragraph 4 ---
    # Errors: teh (#7), straight quotes (#14), lowercase start "this" (#11)
    doc.add_paragraph(
        'The total budget allocation for Q3 was $2.4 million. teh breakdown '
        'is as follows: digital advertising at 45%, content creation at 30%, '
        'and events at 25%. "this represents a strategic shift," noted CFO '
        'Robert Tanaka during the board meeting last Tuesday.'
    )

    # --- Paragraph 5 ---
    # Errors: (c) (#3), lowercase start "the" (#12)
    doc.add_paragraph(
        'All materials produced this quarter carry the updated (c) notice as '
        'required by corporate policy. the design team ensured brand consistency '
        'across every deliverable, from social media graphics to print collateral.'
    )

    # Subheading
    doc.add_heading('Next Steps', level=2)

    # --- Paragraph 6 ---
    # Errors: teh (#8), (c) (#4), straight quotes (#15)
    doc.add_paragraph(
        'Looking ahead to Q4, teh department plans to increase investment in '
        'video content and influencer partnerships. The updated brand guidelines, '
        'including the revised (c) notice format, will be distributed to all '
        'regional offices by October 15. "Our goal is to double engagement '
        'rates by year-end," stated Elena Vasquez at the strategy summit.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
