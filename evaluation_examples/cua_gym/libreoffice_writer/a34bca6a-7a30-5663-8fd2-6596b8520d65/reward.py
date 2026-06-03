"""
Reward Script: Verify AutoCorrect application on document
Task ID: writer_frd_054
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.30): (c) replaced with copyright symbol - 4 instances
  - Component 2 (0.25): 'teh' replaced with 'the' - 4 instances
  - Component 3 (0.25): Lowercase sentence/paragraph starts capitalized
  - Component 4 (0.20): Straight quotes replaced with smart quotes - 3 pairs
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_054'


def persist_app_state(domain: str):
    """Save any unsaved changes in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraph texts for analysis
    all_texts = [para.text for para in doc.paragraphs]

    # Component 1: (c) replaced with copyright symbol (0.30 points)
    # Initial has 4 instances of '(c)' in P1, P3, P6, P8
    # Golden has all replaced with the copyright symbol
    try:
        copyright_sym = '\u00a9'  # copyright symbol
        remaining_c = 0
        copyright_found = 0
        for text in all_texts:
            remaining_c += text.count('(c)')
            copyright_found += text.count(copyright_sym)

        if remaining_c == 0 and copyright_found >= 4:
            print(f"PASS: Component 1 — All (c) replaced with copyright symbol. Found {copyright_found} copyright symbols, 0 remaining (c). (0.30 pts)")
            total_score += 0.30
        elif remaining_c == 0 and copyright_found > 0:
            # All (c) gone but maybe not all converted to symbol (some could be removed)
            partial = 0.30 * (copyright_found / 4.0)
            print(f"PARTIAL: Component 1 — No (c) remaining, but only {copyright_found}/4 copyright symbols found. ({partial:.2f} pts)")
            total_score += min(partial, 0.30)
        elif remaining_c < 4 and copyright_found > 0:
            # Some converted
            converted = 4 - remaining_c
            partial = 0.30 * (converted / 4.0)
            print(f"PARTIAL: Component 1 — {converted}/4 (c) converted. {remaining_c} still remain. ({partial:.2f} pts)")
            total_score += min(partial, 0.30)
        else:
            print(f"FAIL: Component 1 — Found {remaining_c} remaining (c), {copyright_found} copyright symbols")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 'teh' replaced with 'the' (0.25 points)
    # Initial has 4 instances of 'teh ' in P1, P2, P5, P8
    try:
        teh_count = 0
        for text in all_texts:
            # Count 'teh ' occurrences (word boundary: space after, or case-insensitive start)
            import re
            teh_count += len(re.findall(r'\bteh\b', text, re.IGNORECASE))

        if teh_count == 0:
            print(f"PASS: Component 2 — All 'teh' typos corrected. (0.25 pts)")
            total_score += 0.25
        elif teh_count < 4:
            corrected = 4 - teh_count
            partial = 0.25 * (corrected / 4.0)
            print(f"PARTIAL: Component 2 — {corrected}/4 'teh' corrected, {teh_count} remaining. ({partial:.2f} pts)")
            total_score += min(partial, 0.25)
        else:
            print(f"FAIL: Component 2 — Found {teh_count} 'teh' instances still remaining")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Lowercase sentence/paragraph starts capitalized (0.25 points)
    # Initial errors:
    #   P1 starts with lowercase 'teh' -> 'The' (already captured by teh fix, but capitalization matters)
    #   P2 starts with lowercase 'according' -> 'According'
    #   P3 has '. we' -> '. We'
    #   P5 has '. teh' -> '. The' (already captured by teh fix)
    #   P6 has '. the' -> '. The'
    # We check: paragraph starts should be capitalized, and after '. ' should be capitalized
    try:
        import re
        capitalization_errors = 0

        # Check paragraph starts (only body paragraphs, skip headings)
        body_paras = [p for p in doc.paragraphs if p.style and 'Heading' not in p.style.name and p.text.strip()]
        for para in body_paras:
            text = para.text.strip()
            if text and text[0].islower():
                capitalization_errors += 1

        # Check after '. ' within paragraphs
        for text in all_texts:
            # Find lowercase letters after '. '
            matches = re.findall(r'\. ([a-z])', text)
            capitalization_errors += len(matches)

        if capitalization_errors == 0:
            print(f"PASS: Component 3 — All sentence/paragraph starts properly capitalized. (0.25 pts)")
            total_score += 0.25
        elif capitalization_errors <= 2:
            partial = 0.25 * 0.5
            print(f"PARTIAL: Component 3 — {capitalization_errors} capitalization errors remain. ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — {capitalization_errors} capitalization errors found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Straight quotes replaced with smart quotes (0.20 points)
    # Initial has straight double quotes " (U+0022) in P2, P5, P8 (3 pairs = 6 instances)
    # Golden has smart quotes (U+201C opening, U+201D closing)
    try:
        straight_double = 0
        smart_open = 0
        smart_close = 0
        for text in all_texts:
            straight_double += text.count('"')  # U+0022
            smart_open += text.count('\u201c')   # left double quotation mark
            smart_close += text.count('\u201d')   # right double quotation mark

        if straight_double == 0 and smart_open >= 3 and smart_close >= 3:
            print(f"PASS: Component 4 — All straight quotes replaced with smart quotes. Found {smart_open} opening, {smart_close} closing smart quotes. (0.20 pts)")
            total_score += 0.20
        elif straight_double == 0 and (smart_open > 0 or smart_close > 0):
            # Straight quotes gone, but not all converted to smart quotes
            partial = 0.20 * min(smart_open + smart_close, 6) / 6.0
            print(f"PARTIAL: Component 4 — No straight quotes, but only {smart_open}+{smart_close} smart quotes found. ({partial:.2f} pts)")
            total_score += min(partial, 0.20)
        elif straight_double < 6:
            converted = 6 - straight_double
            partial = 0.20 * (converted / 6.0)
            print(f"PARTIAL: Component 4 — {converted}/6 straight quotes converted. {straight_double} remaining. ({partial:.2f} pts)")
            total_score += min(partial, 0.20)
        else:
            print(f"FAIL: Component 4 — Found {straight_double} straight quotes, {smart_open} smart open, {smart_close} smart close")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
