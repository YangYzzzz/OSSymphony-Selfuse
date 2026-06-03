"""
Initial Setup: Disable AutoCorrect capitalize first letter of every sentence
Task ID: writer_frd_046
Domain: libreoffice_writer

Creates a simple Writer document and opens LibreOffice Writer with
default AutoCorrect settings (CapitalAtStartSentence enabled).
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_046'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
REGISTRY = f'{WORKDIR}/.config/libreoffice/4/user/registrymodifications.xcu'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def ensure_autocorrect_default():
    """
    Ensure CapitalAtStartSentence is at default (true) by removing any
    override from registrymodifications.xcu. The default in the schema is true,
    so we just need to make sure there's no explicit false override.
    """
    if not os.path.exists(REGISTRY):
        return

    with open(REGISTRY, 'r', encoding='utf-8') as f:
        content = f.read()

    # Remove any existing CapitalAtStartSentence override
    import re
    # Pattern matches the full <item> element for this setting
    pattern = r'<item oor:path="/org\.openoffice\.Office\.Common/AutoCorrect"><prop oor:name="CapitalAtStartSentence"[^/]*/></item>\n?'
    new_content = re.sub(pattern, '', content)

    if new_content != content:
        with open(REGISTRY, 'w', encoding='utf-8') as f:
            f.write(new_content)
        print('Removed any CapitalAtStartSentence override (default=true)')
    else:
        print('CapitalAtStartSentence is already at default (true)')


def create_initial():
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Add a simple document with realistic content
    doc.add_heading('Quarterly Sales Report', level=1)

    para = doc.add_paragraph()
    run = para.add_run('This report summarizes the quarterly sales performance across all regional offices. ')
    run.font.size = Pt(11)
    run = para.add_run('the data has been compiled from individual branch submissions and verified by the finance team.')
    run.font.size = Pt(11)

    doc.add_heading('Key Highlights', level=2)

    doc.add_paragraph('Total revenue increased by 12% compared to the previous quarter.', style='List Bullet')
    doc.add_paragraph('The Western region exceeded its target by $45,230.', style='List Bullet')
    doc.add_paragraph('Customer acquisition cost decreased from $127 to $98 per customer.', style='List Bullet')

    doc.add_heading('Regional Breakdown', level=2)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['Region', 'Revenue', 'Target', 'Variance']
    for col, h in enumerate(headers):
        table.cell(0, col).text = h

    data = [
        ['Northeast', '$234,500', '$220,000', '+$14,500'],
        ['Southeast', '$198,750', '$200,000', '-$1,250'],
        ['Western', '$265,230', '$220,000', '+$45,230'],
        ['Central', '$187,600', '$190,000', '-$2,400'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_paragraph('For questions about this report, contact the finance department at finance@acmecorp.com.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')


def main():
    # Kill any existing LibreOffice instances
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    time.sleep(1)

    # Ensure default AutoCorrect setting
    ensure_autocorrect_default()

    # Create the document
    create_initial()

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
