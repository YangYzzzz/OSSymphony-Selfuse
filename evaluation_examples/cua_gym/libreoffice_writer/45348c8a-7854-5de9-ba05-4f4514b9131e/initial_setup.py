"""
Initial Setup: Mail merge welcome letter with 8 new hire records
Task ID: writer_mt_008
Domain: libreoffice_writer

Creates:
  - /home/user/writer_mt_008.docx — Welcome letter template with merge fields
  - /home/user/NewHires.csv — Data source with 8 records
  - ~/Desktop/Merged_Letters/ — Empty output folder
  - Opens the document in LibreOffice Writer
"""

import csv
import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
CSV_PATH = f'{WORKDIR}/NewHires.csv'
MERGED_DIR = os.path.expanduser('~/Desktop/Merged_Letters')


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# ── 8 realistic new hire records ──────────────────────────────────

NEW_HIRES = [
    {"FirstName": "Anika", "LastName": "Patel", "Department": "Engineering"},
    {"FirstName": "Marcus", "LastName": "Johnson", "Department": "Marketing"},
    {"FirstName": "Lena", "LastName": "Kowalski", "Department": "Finance"},
    {"FirstName": "David", "LastName": "Chen", "Department": "Product"},
    {"FirstName": "Sofia", "LastName": "Ramirez", "Department": "Human Resources"},
    {"FirstName": "James", "LastName": "Okafor", "Department": "Engineering"},
    {"FirstName": "Yuki", "LastName": "Tanaka", "Department": "Design"},
    {"FirstName": "Rachel", "LastName": "Nguyen", "Department": "Sales"},
]


def create_csv_datasource():
    """Write the NewHires CSV data source."""
    with open(CSV_PATH, 'w', newline='') as f:
        writer = csv.DictWriter(f, fieldnames=["FirstName", "LastName", "Department"])
        writer.writeheader()
        writer.writerows(NEW_HIRES)
    print(f'Data source created: {CSV_PATH}')


def create_template():
    """Build the welcome letter template with merge fields."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.top_margin = Inches(1.2)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ── Company header ──
    header_para = doc.add_paragraph()
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header_para.add_run("Apex Global Solutions")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = sub_para.add_run("Where Innovation Meets Excellence")
    run2.italic = True
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # ── Date line ──
    date_para = doc.add_paragraph("April 1, 2026")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in date_para.runs:
        run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # ── Greeting with merge fields ──
    greeting = doc.add_paragraph()
    r = greeting.add_run("Dear <FirstName> <LastName>,")
    r.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # ── Body paragraphs ──
    body1 = doc.add_paragraph()
    r = body1.add_run(
        "Welcome to Apex Global Solutions! We are thrilled to have you join our "
        "<Department> team. Your skills and experience will be a valuable asset as "
        "we continue to push the boundaries of innovation in our industry."
    )
    r.font.size = Pt(11)
    body1.paragraph_format.space_after = Pt(8)

    body2 = doc.add_paragraph()
    r = body2.add_run(
        "Your onboarding session is scheduled for your first day. Please arrive at "
        "the main lobby by 9:00 AM, where a member of our People Operations team will "
        "greet you and walk you through the day's activities, including a tour of the "
        "office, introductions to your team, and setup of your workstation and accounts."
    )
    r.font.size = Pt(11)
    body2.paragraph_format.space_after = Pt(8)

    body3 = doc.add_paragraph()
    r = body3.add_run(
        "Before your start date, please review the attached Employee Handbook and "
        "complete the pre-boarding checklist sent to your email. If you have any "
        "questions, do not hesitate to reach out to your hiring manager or the HR team."
    )
    r.font.size = Pt(11)
    body3.paragraph_format.space_after = Pt(8)

    body4 = doc.add_paragraph()
    r = body4.add_run(
        "We look forward to seeing you on your first day and are confident that your "
        "contributions will help us achieve great things together."
    )
    r.font.size = Pt(11)
    body4.paragraph_format.space_after = Pt(16)

    # ── Closing ──
    closing = doc.add_paragraph()
    r = closing.add_run("Warm regards,")
    r.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    sig = doc.add_paragraph()
    r = sig.add_run("Jennifer Martinez")
    r.bold = True
    r.font.size = Pt(11)

    title_para = doc.add_paragraph()
    r = title_para.add_run("Vice President of People Operations")
    r.font.size = Pt(11)

    company_para = doc.add_paragraph()
    r = company_para.add_run("Apex Global Solutions")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.save(OUTPUT)
    print(f'Template created: {OUTPUT}')


def create_output_folder():
    """Create the empty Merged_Letters folder on the Desktop."""
    os.makedirs(MERGED_DIR, exist_ok=True)
    print(f'Output folder created: {MERGED_DIR}')


def main():
    create_csv_datasource()
    create_template()
    create_output_folder()

    # Open the template in LibreOffice Writer (GUI-ready)
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
