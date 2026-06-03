"""
Initial Setup: Apply Heading 3 style to four sub-sub-headings in a research report
Task ID: writer_struct_030
Domain: libreoffice_writer

Creates experiment_report.docx with 'Sampling Method', 'Survey Design',
'Control Group Selection', and 'Data Normalization' as Default Paragraph Style
(NOT Heading 3 — the agent must apply that style).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'experiment_report'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # ---- Title ----
    title = doc.add_heading('Experimental Research Report: Behavioral Analytics in Consumer Decision-Making', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---- Abstract ----
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This report presents a comprehensive investigation into consumer behavioral patterns '
        'within digital retail environments. The study draws on quantitative survey data collected '
        'from 1,240 participants across six metropolitan regions. Findings indicate significant '
        'correlations between interface design complexity and purchase conversion rates, with '
        'implications for UX-driven product strategy. Statistical models were validated using '
        'cross-sectional regression and bootstrapped confidence intervals.'
    )

    # ---- Introduction ----
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'The rapid expansion of e-commerce platforms over the past decade has generated an '
        'unprecedented volume of behavioral data. Understanding the mechanisms that drive '
        'consumer decision-making is central to marketing science, human-computer interaction, '
        'and operations management. The present study addresses a gap in the literature by '
        'examining how specific interface features interact with user cognitive load to shape '
        'purchasing behavior.'
    )
    doc.add_paragraph(
        'Previous research has established foundational relationships between usability and '
        'conversion, yet few studies have applied longitudinal tracking combined with controlled '
        'experimental conditions. This report details our multi-phase methodology, including '
        'participant sampling, survey design rationale, control group management, and the '
        'normalization procedures applied to the collected data.'
    )

    # ---- Literature Review ----
    doc.add_heading('Literature Review', level=1)
    doc.add_paragraph(
        'Foundational work by Nielsen and Molich (1990) established heuristic evaluation as a '
        'cornerstone of usability assessment. Subsequent contributions from Shneiderman (1992) '
        'extended these principles to direct-manipulation interfaces. More recently, Fogg (2003) '
        'introduced the concept of persuasive technology, arguing that interface design carries '
        'implicit motivational affordances.'
    )
    doc.add_paragraph(
        'In the domain of consumer psychology, Kahneman\'s dual-process theory (2011) has been '
        'extensively applied to model impulsive vs. deliberative purchasing decisions. Behavioral '
        'economists such as Ariely (2008) have demonstrated that presentation framing—including '
        'default options, anchoring prices, and scarcity signals—strongly influences consumer '
        'choice architecture. This study integrates both traditions to analyze interaction effects '
        'between cognitive load and interface persuasion features.'
    )

    # ---- Research Methods ----
    doc.add_heading('Research Methods', level=2)
    doc.add_paragraph(
        'The methodology adopted in this study was designed to ensure internal validity while '
        'preserving ecological validity across multiple testing environments. The following '
        'subsections describe each component of the research design.'
    )

    # --- Sub-sub-headings styled as Default Paragraph Style (NOT Heading 3) ---
    # The agent must change these to Heading 3

    # Sampling Method
    p_sm = doc.add_paragraph('Sampling Method')
    # Style is already 'Normal' / Default Paragraph Style — intentionally NOT Heading 3
    doc.add_paragraph(
        'Participants were recruited using stratified random sampling across six cities: '
        'New York, Los Angeles, Chicago, Houston, Phoenix, and Philadelphia. Quotas were '
        'set for age (18–34, 35–54, 55+), gender, and household income bracket to ensure '
        'demographic representativeness. A total of 1,240 individuals completed the full '
        'study protocol. Recruitment was conducted through a third-party panel provider '
        'with ISO-certified data quality controls. Dropout rates were managed through '
        'staged incentive structures, resulting in a completion rate of 87.3%.'
    )

    # Survey Design
    p_sd = doc.add_paragraph('Survey Design')
    doc.add_paragraph(
        'The survey instrument comprised 48 items organized across five thematic modules: '
        'general demographics, technology adoption attitudes, prior online purchase history, '
        'interface usability perception, and post-task decision confidence. Items were '
        'developed through iterative expert review involving three UX researchers and two '
        'consumer psychologists. A 7-point Likert scale was used for all attitudinal '
        'measures. Pilot testing with 60 participants confirmed internal consistency '
        '(Cronbach\'s alpha = 0.84) and convergent validity with established scales.'
    )

    # Control Group Selection
    p_cg = doc.add_paragraph('Control Group Selection')
    doc.add_paragraph(
        'The control group (n=310) was selected to mirror the experimental group on all '
        'measured covariates. Matching was performed using propensity score matching on '
        'age, gender, prior e-commerce usage frequency, and self-reported technology '
        'proficiency. Balance diagnostics confirmed standardized mean differences below '
        '0.10 for all covariates post-matching. Control participants interacted with a '
        'baseline interface version stripped of all persuasive design elements, including '
        'urgency messaging, social proof indicators, and dynamic pricing cues.'
    )

    # Data Normalization
    p_dn = doc.add_paragraph('Data Normalization')
    doc.add_paragraph(
        'Raw response data were subjected to a multi-step normalization pipeline prior '
        'to analysis. Outliers were identified using the IQR method (1.5× fence) and '
        'verified through manual inspection of response patterns. Z-score normalization '
        'was applied to continuous variables to eliminate scale effects. Ordinal Likert '
        'responses were treated as interval-level data following established conventions '
        'in behavioral research (Norman, 2010). Missing values were imputed using '
        'multivariate imputation by chained equations (MICE) with five imputation cycles.'
    )

    # ---- Results ----
    doc.add_heading('Results', level=1)
    doc.add_paragraph(
        'Descriptive statistics revealed that 68.4% of experimental group participants '
        'completed a simulated purchase task compared to 51.2% in the control group '
        '(χ² = 41.7, df=1, p<0.001). Regression analysis identified interface simplicity '
        'score (β=0.42, p<0.001) and scarcity signal presence (β=0.29, p=0.003) as the '
        'strongest predictors of conversion intent.'
    )
    doc.add_paragraph(
        'Moderating analysis using interaction terms confirmed that the effect of '
        'persuasive design was significantly attenuated for high-proficiency users '
        '(interaction β=−0.18, p=0.021). No significant gender differences were observed '
        'after controlling for technology adoption attitudes. Age cohort analysis revealed '
        'the 35–54 group showed the highest susceptibility to urgency messaging (β=0.51).'
    )

    # ---- Discussion ----
    doc.add_heading('Discussion', level=1)
    doc.add_paragraph(
        'The results support the hypothesis that persuasive interface design significantly '
        'increases purchase conversion, consistent with Fogg\'s (2003) theoretical framework. '
        'However, the moderating role of technology proficiency suggests that experienced '
        'users develop resistance to persuasion cues—a finding with important implications '
        'for adaptive interface design strategies targeting heterogeneous user bases.'
    )
    doc.add_paragraph(
        'The strong performance of scarcity signals replicates findings from prior laboratory '
        'studies (Cialdini, 1984; Aggarwal et al., 2011) in a more ecologically valid online '
        'context. The null result for gender differences diverges from some prior work, '
        'possibly reflecting cohort effects or the equalization of digital familiarity '
        'across genders in the current sample demographics.'
    )

    # ---- Conclusion ----
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'This study provides empirical evidence that targeted persuasive interface elements '
        'reliably increase consumer conversion rates in e-commerce contexts, while also '
        'identifying boundary conditions related to user expertise. Future research should '
        'examine longitudinal effects and the potential for personalization-based delivery '
        'of persuasive cues to optimize effectiveness without eroding user trust.'
    )
    doc.add_paragraph(
        'Methodological contributions include the validated sampling protocol, the '
        'propensity-matched control group design, and the normalization pipeline applicable '
        'to large-scale behavioral survey datasets. These frameworks are made available in '
        'the supplementary appendix for replication and extension by other researchers.'
    )

    # ---- References ----
    doc.add_heading('References', level=1)
    refs = [
        'Aggarwal, P., Jun, S. Y., & Huh, J. H. (2011). Scarcity messages. Journal of Advertising, 40(3), 19–30.',
        'Ariely, D. (2008). Predictably Irrational. HarperCollins.',
        'Cialdini, R. B. (1984). Influence: The Psychology of Persuasion. Harper Business.',
        'Fogg, B. J. (2003). Persuasive Technology. Morgan Kaufmann.',
        'Kahneman, D. (2011). Thinking, Fast and Slow. Farrar, Straus and Giroux.',
        'Nielsen, J., & Molich, R. (1990). Heuristic evaluation of user interfaces. CHI\'90 Proceedings.',
        'Norman, G. (2010). Likert scales, levels of measurement and the "laws" of statistics. Advances in Health Sciences Education, 15(5), 625–632.',
        'Shneiderman, B. (1992). Designing the User Interface. Addison-Wesley.',
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='List Bullet')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
