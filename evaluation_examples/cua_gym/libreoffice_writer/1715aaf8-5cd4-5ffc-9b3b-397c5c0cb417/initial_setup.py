"""
Initial Setup: Add AutoCorrect exception for 'vs.' abbreviation
Task ID: writer_frd_056
Domain: libreoffice_writer

Creates a user-level autocorrect file where 'vs.' is NOT in the sentence
exception list, so typing 'plaintiff vs. defendant' auto-capitalizes 'defendant'.
Also creates a sample document and opens it in LibreOffice Writer.
"""

import os
import shlex
import shutil
import subprocess
import tempfile
import time
import zipfile

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_056'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

SYSTEM_DAT = '/usr/lib/libreoffice/share/autocorr/acor_en-US.dat'
USER_AUTOCORR_DIR = os.path.expanduser('~/.config/libreoffice/4/user/autocorr')
USER_DAT = os.path.join(USER_AUTOCORR_DIR, 'acor_en-US.dat')


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_user_autocorr_without_vs():
    """
    Copy system acor_en-US.dat to user profile, but remove 'vs.' from
    SentenceExceptList.xml so that 'vs.' triggers sentence capitalization.
    """
    os.makedirs(USER_AUTOCORR_DIR, exist_ok=True)

    # Read system .dat (ZIP), modify SentenceExceptList.xml, write to user .dat
    with zipfile.ZipFile(SYSTEM_DAT, 'r') as zin:
        with zipfile.ZipFile(USER_DAT, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                data = zin.read(item)
                if item == 'SentenceExceptList.xml':
                    content = data.decode('utf-8')
                    # Remove the 'vs.' entry
                    content = content.replace(
                        '<block-list:block block-list:abbreviated-name="vs."/>',
                        ''
                    )
                    data = content.encode('utf-8')
                zout.writestr(item, data)

    print(f'User autocorrect file created at {USER_DAT} (without vs. exception)')


def create_sample_document():
    """
    Create a sample legal document where the agent will test autocorrect behavior.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Title
    heading = doc.add_heading('Legal Brief - Motion for Summary Judgment', level=1)

    # Introduction paragraph
    doc.add_paragraph(
        'This memorandum is submitted on behalf of the plaintiff in the matter of '
        'Thompson Industries, Inc. v. Meridian Supply Co., Case No. 2025-CV-04891.'
    )

    doc.add_heading('Statement of Facts', level=2)
    doc.add_paragraph(
        'On March 15, 2025, the plaintiff entered into a supply agreement with the '
        'defendant for the delivery of industrial components. The contract specified '
        'delivery within 30 business days of each purchase order.'
    )
    doc.add_paragraph(
        'Between April and September 2025, the defendant failed to meet delivery '
        'deadlines on seven separate occasions, causing significant production delays '
        'at the plaintiff\'s manufacturing facility in Portland, Oregon.'
    )

    doc.add_heading('Legal Standard', level=2)
    doc.add_paragraph(
        'Summary judgment is appropriate when the movant shows that there is no genuine '
        'dispute as to any material fact and the movant is entitled to judgment as a '
        'matter of law. Fed. R. Civ. P. 56(a).'
    )

    doc.add_heading('Argument', level=2)
    doc.add_paragraph(
        'The undisputed facts demonstrate that the defendant breached the supply '
        'agreement repeatedly. The evidence clearly shows a pattern of delayed '
        'deliveries that directly impacted the plaintiff\'s operations.'
    )

    # This paragraph is where the user would notice the autocorrect issue
    doc.add_paragraph(
        'In evaluating contract performance standards, courts have consistently held '
        'that repeated failures constitute a material breach. The analysis of '
        'plaintiff vs. defendant obligations supports this conclusion.'
    )

    doc.add_heading('Damages', level=2)
    doc.add_paragraph(
        'The plaintiff seeks compensatory damages in the amount of $287,450.00, '
        'representing lost production revenue ($198,200.00), expedited shipping costs '
        'for alternative suppliers ($52,750.00), and administrative overhead ($36,500.00).'
    )

    doc.add_heading('Conclusion', level=2)
    doc.add_paragraph(
        'For the foregoing reasons, the plaintiff respectfully requests that this Court '
        'grant summary judgment in its favor on all counts.'
    )

    doc.save(OUTPUT)
    print(f'Sample document created: {OUTPUT}')


def main():
    # Kill any running LibreOffice instances first
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    time.sleep(1)

    # Step 1: Create the user autocorrect file without 'vs.'
    create_user_autocorr_without_vs()

    # Step 2: Create the sample document
    create_sample_document()

    # Step 3: Open LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
