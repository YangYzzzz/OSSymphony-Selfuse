"""
Initial Setup: Software Manual with screenshots and tables, no captions or lists
Task ID: writer_pd_023
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
import struct
import zlib

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_023'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def make_placeholder_png(width=640, height=400, label="Screenshot", color=(200, 210, 230)):
    """Create a minimal PNG in memory (no PIL needed on VM)."""
    r, g, b = color
    raw_rows = []
    for y in range(height):
        row = b'\x00'  # filter byte
        for x in range(width):
            row += bytes([r, g, b])
        raw_rows.append(row)
    raw_data = b''.join(raw_rows)

    def chunk(chunk_type, data):
        c = chunk_type + data
        return struct.pack('>I', len(data)) + c + struct.pack('>I', zlib.crc32(c) & 0xffffffff)

    ihdr_data = struct.pack('>IIBBBBB', width, height, 8, 2, 0, 0, 0)
    compressed = zlib.compress(raw_data)
    png = b'\x89PNG\r\n\x1a\n'
    png += chunk(b'IHDR', ihdr_data)
    png += chunk(b'IDAT', compressed)
    png += chunk(b'IEND', b'')
    return png


def launch_gui(command: str, delay_sec: float = 1.0):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Save placeholder images to disk
    screenshot_colors = [
        (180, 200, 230), (200, 220, 200), (230, 210, 190),
        (210, 190, 220), (190, 220, 220), (220, 200, 200),
    ]
    screenshot_labels = [
        "Dashboard Overview", "User Management Panel", "Configuration Settings",
        "Report Generation View", "API Integration Console", "System Monitoring Dashboard"
    ]
    img_paths = []
    for i in range(6):
        path = f'/tmp/screenshot_{i+1}.png'
        with open(path, 'wb') as f:
            f.write(make_placeholder_png(640, 400, screenshot_labels[i], screenshot_colors[i]))
        img_paths.append(path)

    # ============ PAGE 1: Title Page ============
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_heading("CloudSync Pro 3.0", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_heading("Software User Manual", level=1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for _ in range(2):
        doc.add_paragraph("")

    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = info.add_run("Version 3.0.2 | Release Date: March 15, 2025\nPrepared by: Technical Documentation Team\nNovaTech Solutions Inc.")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ============ PAGE 2: Blank (reserved for lists) ============
    doc.add_page_break()
    reserved = doc.add_paragraph("")
    reserved.paragraph_format.space_after = Pt(0)
    # Intentionally left blank for List of Figures and List of Tables

    # ============ PAGE 3: Chapter 1 with Screenshot 1 ============
    doc.add_page_break()
    doc.add_heading("Chapter 1: Getting Started", level=1)
    doc.add_paragraph(
        "CloudSync Pro 3.0 is a comprehensive cloud-based synchronization platform designed for "
        "enterprise teams. This chapter walks you through the initial setup process, including "
        "account creation, workspace configuration, and connecting your first data sources."
    )
    doc.add_paragraph(
        "After logging in for the first time, you will be greeted by the main dashboard. "
        "The dashboard provides an at-a-glance overview of your synchronization status, "
        "recent activity, and system notifications."
    )
    # Screenshot 1
    doc.add_picture(img_paths[0], width=Inches(5.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        "The navigation panel on the left provides quick access to all major modules. "
        "You can customize the dashboard layout by clicking the gear icon in the top-right corner."
    )

    # ============ PAGE 4: Table 1 ============
    doc.add_page_break()
    doc.add_heading("1.1 System Requirements", level=2)
    doc.add_paragraph(
        "Before installing CloudSync Pro, verify that your system meets the following minimum requirements:"
    )
    # Table 1: System Requirements
    table1 = doc.add_table(rows=7, cols=3)
    table1.style = "Table Grid"
    headers1 = ["Component", "Minimum Requirement", "Recommended"]
    for i, h in enumerate(headers1):
        cell = table1.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data1 = [
        ["Operating System", "Windows 10 / macOS 12 / Ubuntu 20.04", "Windows 11 / macOS 14 / Ubuntu 22.04"],
        ["Processor", "Intel Core i5 (8th Gen)", "Intel Core i7 (12th Gen) or AMD Ryzen 7"],
        ["RAM", "8 GB", "16 GB"],
        ["Storage", "500 MB available", "2 GB SSD"],
        ["Network", "10 Mbps broadband", "100 Mbps+ fiber"],
        ["Browser", "Chrome 90+, Firefox 88+", "Chrome 120+, Edge 120+"],
    ]
    for r, row_data in enumerate(data1, 1):
        for c, val in enumerate(row_data):
            table1.cell(r, c).text = val

    doc.add_paragraph(
        "For optimal performance in enterprise deployments with more than 500 concurrent users, "
        "we recommend deploying on dedicated infrastructure with the recommended specifications."
    )

    # ============ PAGE 5: More text + Screenshot 2 ============
    doc.add_page_break()
    doc.add_heading("Chapter 2: User Management", level=1)
    doc.add_paragraph(
        "CloudSync Pro supports role-based access control (RBAC) with customizable permission levels. "
        "Administrators can create user groups, assign roles, and manage authentication policies "
        "from the User Management panel."
    )
    doc.add_paragraph(
        "The user management interface provides a centralized view of all registered accounts, "
        "their roles, last login timestamps, and current synchronization quotas."
    )
    # Screenshot 2
    doc.add_picture(img_paths[1], width=Inches(5.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        "To add a new user, click the 'Add User' button and fill in the required fields. "
        "The system will send an activation email to the new user's registered address."
    )

    # ============ PAGE 6: Additional content ============
    doc.add_page_break()
    doc.add_heading("2.1 Role Definitions", level=2)
    doc.add_paragraph(
        "CloudSync Pro ships with four predefined roles: Administrator, Manager, Editor, and Viewer. "
        "Each role has a distinct set of permissions that control access to synchronization jobs, "
        "configuration settings, and reporting dashboards."
    )
    doc.add_paragraph(
        "Custom roles can be created by cloning an existing role and modifying individual permissions. "
        "Role changes take effect immediately for all users assigned to that role."
    )
    doc.add_paragraph(
        "Security best practices recommend following the principle of least privilege when assigning roles. "
        "Regular audits of user permissions help ensure compliance with organizational security policies."
    )

    # ============ PAGE 7: Screenshot 3 ============
    doc.add_page_break()
    doc.add_heading("Chapter 3: Configuration", level=1)
    doc.add_paragraph(
        "The configuration module allows administrators to fine-tune synchronization behavior, "
        "set up automated schedules, and configure notification preferences. All settings are "
        "organized into logical categories for easy navigation."
    )
    # Screenshot 3
    doc.add_picture(img_paths[2], width=Inches(5.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        "Changes to synchronization settings require a confirmation step before they take effect. "
        "A preview of the impact is shown, including the number of affected sync jobs."
    )

    # ============ PAGE 8: Table 2 ============
    doc.add_page_break()
    doc.add_heading("3.1 Synchronization Parameters", level=2)
    doc.add_paragraph(
        "The following table lists the key synchronization parameters and their default values:"
    )
    # Table 2: Sync Parameters
    table2 = doc.add_table(rows=8, cols=4)
    table2.style = "Table Grid"
    headers2 = ["Parameter", "Default Value", "Range", "Description"]
    for i, h in enumerate(headers2):
        cell = table2.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data2 = [
        ["Sync Interval", "15 minutes", "1-1440 min", "Time between automatic syncs"],
        ["Batch Size", "500 records", "100-10000", "Records processed per batch"],
        ["Retry Count", "3", "0-10", "Number of retry attempts on failure"],
        ["Timeout", "120 seconds", "30-600 sec", "Max wait time per operation"],
        ["Compression", "Enabled", "On/Off", "Compress data during transfer"],
        ["Encryption", "AES-256", "AES-128/256", "Encryption algorithm for transit"],
        ["Log Level", "INFO", "DEBUG-FATAL", "Verbosity of sync logs"],
    ]
    for r, row_data in enumerate(data2, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    # ============ PAGE 9: Screenshot 4 ============
    doc.add_page_break()
    doc.add_heading("Chapter 4: Reports and Analytics", level=1)
    doc.add_paragraph(
        "CloudSync Pro includes a comprehensive reporting engine that tracks synchronization "
        "performance, error rates, and data throughput. Reports can be generated on demand "
        "or scheduled for automatic delivery via email."
    )
    # Screenshot 4
    doc.add_picture(img_paths[3], width=Inches(5.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        "The analytics dashboard provides interactive charts and drill-down capabilities. "
        "Click on any data point to view detailed synchronization logs for that time period."
    )

    # ============ PAGE 10: More content ============
    doc.add_page_break()
    doc.add_heading("4.1 Performance Metrics", level=2)
    doc.add_paragraph(
        "Key performance indicators (KPIs) are tracked in real-time and displayed on the "
        "analytics dashboard. These include average sync duration, success rate, data volume "
        "processed, and active connection count."
    )
    doc.add_paragraph(
        "Historical performance data is retained for 90 days by default. Enterprise customers "
        "can extend this retention period to 365 days through the advanced configuration panel."
    )
    doc.add_paragraph(
        "Anomaly detection algorithms automatically flag unusual patterns in sync behavior, "
        "such as sudden spikes in error rates or unexpected drops in throughput."
    )

    # ============ PAGE 11: Table 3 ============
    doc.add_page_break()
    doc.add_heading("4.2 Error Code Reference", level=2)
    doc.add_paragraph(
        "The following table provides a reference for common synchronization error codes:"
    )
    # Table 3: Error Codes
    table3 = doc.add_table(rows=9, cols=4)
    table3.style = "Table Grid"
    headers3 = ["Error Code", "Severity", "Description", "Resolution"]
    for i, h in enumerate(headers3):
        cell = table3.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data3 = [
        ["E1001", "Critical", "Authentication failure", "Verify API credentials"],
        ["E1002", "High", "Connection timeout", "Check network connectivity"],
        ["E1003", "Medium", "Schema mismatch", "Update field mappings"],
        ["E1004", "Low", "Duplicate record detected", "Review dedup rules"],
        ["E1005", "Critical", "Storage quota exceeded", "Increase storage allocation"],
        ["E1006", "High", "Rate limit exceeded", "Reduce sync frequency"],
        ["E1007", "Medium", "Partial sync failure", "Review failed records log"],
        ["E1008", "Low", "Deprecated API version", "Update integration endpoint"],
    ]
    for r, row_data in enumerate(data3, 1):
        for c, val in enumerate(row_data):
            table3.cell(r, c).text = val

    # ============ PAGE 12: Screenshot 5 ============
    doc.add_page_break()
    doc.add_heading("Chapter 5: API Integration", level=1)
    doc.add_paragraph(
        "CloudSync Pro exposes a RESTful API for programmatic access to all synchronization "
        "features. The API supports JSON and XML payloads, with OAuth 2.0 authentication."
    )
    doc.add_paragraph(
        "The API integration console provides a live sandbox environment where developers "
        "can test API calls and review response payloads before deploying to production."
    )
    # Screenshot 5
    doc.add_picture(img_paths[4], width=Inches(5.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        "API rate limits are configured per user role. Administrator accounts have unlimited "
        "API access, while standard accounts are limited to 1000 requests per hour."
    )

    # ============ PAGES 13-14: More content ============
    doc.add_page_break()
    doc.add_heading("5.1 Authentication Flow", level=2)
    doc.add_paragraph(
        "All API requests must include a valid Bearer token in the Authorization header. "
        "Tokens are obtained through the OAuth 2.0 client credentials flow."
    )
    doc.add_paragraph(
        "Access tokens have a default lifetime of 3600 seconds (1 hour). Refresh tokens "
        "can be used to obtain new access tokens without re-authenticating."
    )
    doc.add_heading("5.2 Webhook Configuration", level=2)
    doc.add_paragraph(
        "Webhooks allow external systems to receive real-time notifications about synchronization "
        "events. Supported event types include sync.started, sync.completed, sync.failed, and record.conflict."
    )
    doc.add_paragraph(
        "Webhook endpoints must respond with HTTP 200 within 10 seconds. Failed deliveries "
        "are retried up to 5 times with exponential backoff."
    )
    doc.add_paragraph(
        "Each webhook payload includes a cryptographic signature in the X-CloudSync-Signature "
        "header, allowing receivers to verify the authenticity of the notification."
    )

    # ============ PAGE 15: Screenshot 6 ============
    doc.add_page_break()
    doc.add_heading("Chapter 6: System Monitoring", level=1)
    doc.add_paragraph(
        "The system monitoring module provides real-time visibility into CloudSync Pro's "
        "operational health. Infrastructure metrics, service status, and resource utilization "
        "are displayed on an interactive dashboard."
    )
    # Screenshot 6
    doc.add_picture(img_paths[5], width=Inches(5.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        "Alerts can be configured to notify administrators when metrics exceed defined thresholds. "
        "Integration with popular monitoring platforms like PagerDuty and Opsgenie is available."
    )

    # ============ PAGE 16: Table 4 ============
    doc.add_page_break()
    doc.add_heading("6.1 Service Level Objectives", level=2)
    doc.add_paragraph(
        "CloudSync Pro maintains the following service level objectives for enterprise customers:"
    )
    # Table 4: SLO Table
    table4 = doc.add_table(rows=7, cols=4)
    table4.style = "Table Grid"
    headers4 = ["Metric", "Target", "Measurement Window", "Penalty"]
    for i, h in enumerate(headers4):
        cell = table4.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data4 = [
        ["Uptime", "99.95%", "Monthly", "10% credit per 0.1% below target"],
        ["Sync Latency", "< 500ms (P95)", "Weekly", "5% credit if exceeded"],
        ["Error Rate", "< 0.1%", "Daily", "Service review triggered"],
        ["Data Throughput", "> 10 GB/hour", "Hourly", "Auto-scaling initiated"],
        ["API Response Time", "< 200ms (P99)", "Daily", "Performance investigation"],
        ["Recovery Time", "< 15 minutes", "Per incident", "Post-mortem required"],
    ]
    for r, row_data in enumerate(data4, 1):
        for c, val in enumerate(row_data):
            table4.cell(r, c).text = val

    # ============ PAGES 17-18: Appendix ============
    doc.add_page_break()
    doc.add_heading("Appendix A: Glossary", level=1)
    terms = [
        ("RBAC", "Role-Based Access Control. A method of regulating access to resources based on the roles of individual users within an organization."),
        ("OAuth 2.0", "An open standard for access delegation, commonly used for token-based authentication and authorization."),
        ("Webhook", "A user-defined HTTP callback that is triggered by a specific event, enabling real-time communication between applications."),
        ("SLO", "Service Level Objective. A target value for a service level indicator that is agreed upon between a provider and customer."),
        ("P95/P99", "The 95th or 99th percentile of a distribution, commonly used to measure latency excluding extreme outliers."),
        ("Deduplication", "The process of identifying and eliminating duplicate records to maintain data integrity and reduce storage requirements."),
    ]
    for term, definition in terms:
        p = doc.add_paragraph()
        run_term = p.add_run(term + ": ")
        run_term.bold = True
        p.add_run(definition)

    doc.add_page_break()
    doc.add_heading("Appendix B: Revision History", level=1)
    revision_table = doc.add_table(rows=5, cols=4)
    revision_table.style = "Table Grid"
    rev_headers = ["Version", "Date", "Author", "Changes"]
    for i, h in enumerate(rev_headers):
        cell = revision_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    rev_data = [
        ["3.0.2", "2025-03-15", "Emily Rodriguez", "Updated API integration chapter"],
        ["3.0.1", "2025-01-20", "David Park", "Added error code reference table"],
        ["3.0.0", "2024-11-10", "Sarah Mitchell", "Major release - complete rewrite"],
        ["2.5.0", "2024-06-01", "James Wilson", "Added monitoring module documentation"],
    ]
    for r, row_data in enumerate(rev_data, 1):
        for c, val in enumerate(row_data):
            revision_table.cell(r, c).text = val

    doc.add_paragraph("")
    doc.add_paragraph(
        "Copyright 2025 NovaTech Solutions Inc. All rights reserved. This document may not be "
        "reproduced or distributed without written permission."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Clean up temp images
    for p in img_paths:
        try:
            os.remove(p)
        except OSError:
            pass

    # GUI-ready: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
