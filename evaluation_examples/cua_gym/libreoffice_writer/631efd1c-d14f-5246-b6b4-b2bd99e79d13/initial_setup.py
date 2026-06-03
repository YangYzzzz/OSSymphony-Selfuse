"""
Initial Setup: Create a multi-page company report with [COMPANY] placeholders
Task ID: writer_tm_075
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_075'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Header with [COMPANY] placeholder ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = hp.add_run('[COMPANY]')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Add a thin line under header
    hp2 = header.add_paragraph()
    hp2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = hp2.add_run('Annual Strategic Report 2025')
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # --- Footer with [COMPANY] placeholder ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_f = fp.add_run('[COMPANY]')
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    fp.add_run(' | Confidential | ').font.size = Pt(8)
    # Page number field
    r1 = fp.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)
    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)

    # ===================== PAGE 1 =====================
    # Title
    title = doc.add_heading('Annual Strategic Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # First paragraph with [COMPANY] placeholder
    p1 = doc.add_paragraph()
    p1.add_run('This report presents the comprehensive strategic overview prepared by ')
    run_c1 = p1.add_run('[COMPANY]')
    run_c1.bold = True
    p1.add_run(' for the fiscal year ending December 2025. The document outlines our key '
               'achievements, financial performance, market position, and forward-looking '
               'strategic initiatives designed to maintain our competitive advantage in the '
               'rapidly evolving technology landscape.')

    p1b = doc.add_paragraph()
    p1b.add_run('The information contained herein is intended for internal stakeholders, '
                 'board members, and senior management. All figures have been audited by '
                 'Thornton & Associates LLP and reflect our consolidated operations across '
                 'all regional divisions.')

    # ===================== PAGE 2 =====================
    doc.add_page_break()

    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'The fiscal year 2025 marked a transformative period for our organization. '
        'Revenue grew by 23% year-over-year, reaching $847.3 million, driven primarily '
        'by expansion in the Asia-Pacific region and the successful launch of our '
        'CloudSync Enterprise platform. Operating margins improved to 18.2%, up from '
        '15.7% in FY2024, reflecting our ongoing cost optimization initiatives.'
    )

    doc.add_paragraph(
        'Key highlights of the year include the acquisition of DataStream Analytics, '
        'a strategic partnership with Meridian Systems for joint cloud infrastructure '
        'development, and the opening of new R&D centers in Bangalore and Tel Aviv. '
        'Our workforce expanded to 4,200 employees across 12 countries, with a '
        'voluntary attrition rate of just 8.3%.'
    )

    doc.add_heading('Financial Highlights', level=2)

    # Financial summary table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers_data = ['Metric', 'FY2025', 'FY2024']
    for i, h in enumerate(headers_data):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    fin_data = [
        ['Total Revenue', '$847.3M', '$688.9M'],
        ['Operating Income', '$154.2M', '$108.2M'],
        ['Net Income', '$121.8M', '$89.4M'],
        ['Earnings per Share', '$3.42', '$2.51'],
        ['Free Cash Flow', '$198.7M', '$156.3M'],
    ]
    for r, row_data in enumerate(fin_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()  # spacing

    # ===================== PAGE 3 =====================
    doc.add_page_break()

    doc.add_heading('Market Analysis & Competitive Landscape', level=1)
    doc.add_paragraph(
        'The global enterprise software market reached $672 billion in 2025, growing at '
        'a compound annual rate of 11.2%. Cloud-native solutions continued to gain market '
        'share, with hybrid cloud deployments becoming the predominant architecture for '
        'mid-to-large enterprises.'
    )

    # Page 3 paragraph with [COMPANY] placeholder
    p3 = doc.add_paragraph()
    p3.add_run('In this competitive landscape, ')
    run_c3 = p3.add_run('[COMPANY]')
    run_c3.bold = True
    p3.add_run(' has successfully differentiated itself through its proprietary AI-driven '
               'automation platform, which reduces deployment times by 40% compared to '
               'industry benchmarks. Our Net Promoter Score of 72 places us in the top '
               'quartile of enterprise technology providers.')

    doc.add_paragraph(
        'Primary competitors include Cascade Solutions (market share 15.3%), '
        'Pinnacle Software Group (12.8%), and Horizon Digital (11.1%). Our own market '
        'share grew from 9.4% to 11.6%, positioning us as the fourth-largest provider '
        'in our segment. The acquisition of DataStream Analytics is expected to '
        'add approximately 2.1 percentage points to our market share in FY2026.'
    )

    doc.add_heading('Regional Performance', level=2)
    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'
    reg_headers = ['Region', 'Revenue', 'Growth', 'Key Market']
    for i, h in enumerate(reg_headers):
        cell = table2.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    reg_data = [
        ['North America', '$412.6M', '+18%', 'United States'],
        ['Europe', '$198.3M', '+21%', 'Germany'],
        ['Asia-Pacific', '$187.2M', '+34%', 'Japan'],
        ['Rest of World', '$49.2M', '+27%', 'Brazil'],
    ]
    for r, row_data in enumerate(reg_data, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    # ===================== PAGE 4 =====================
    doc.add_page_break()

    doc.add_heading('Product & Technology Innovation', level=1)
    doc.add_paragraph(
        'Our R&D investment of $127.4 million (15% of revenue) funded several '
        'breakthrough initiatives in FY2025. The CloudSync Enterprise platform, '
        'launched in Q2, has already been adopted by 340 enterprise clients and '
        'contributed $62.8 million in new annual recurring revenue.'
    )

    doc.add_paragraph(
        'The integration of large language models into our customer support automation '
        'suite reduced average resolution time from 4.2 hours to 1.1 hours, while '
        'improving customer satisfaction scores by 15 points. Our patent portfolio '
        'grew to 287 active patents, with 43 new filings in FY2025.'
    )

    doc.add_heading('Technology Roadmap', level=2)
    tech_items = [
        'Q1 2026: Launch of CloudSync 2.0 with multi-cloud orchestration',
        'Q2 2026: AI-powered predictive analytics module',
        'Q3 2026: Zero-trust security framework integration',
        'Q4 2026: Edge computing support for IoT workloads',
    ]
    for item in tech_items:
        doc.add_paragraph(item, style='List Bullet')

    # ===================== PAGE 5 =====================
    doc.add_page_break()

    doc.add_heading('Human Capital & Organizational Development', level=1)
    doc.add_paragraph(
        'Our people strategy remains central to our competitive advantage. In FY2025, '
        'we hired 820 new team members while maintaining our rigorous selection standards. '
        'The average time-to-hire decreased from 42 days to 31 days through process '
        'automation and AI-assisted candidate screening.'
    )

    doc.add_paragraph(
        'Employee engagement, measured through our quarterly pulse surveys, reached '
        'an all-time high of 4.3 out of 5.0. Our investment in learning and development '
        'totaled $8.4 million, with each employee averaging 48 hours of professional '
        'development. Leadership development programs graduated 85 participants who '
        'moved into senior roles.'
    )

    doc.add_heading('Diversity & Inclusion Metrics', level=2)
    table3 = doc.add_table(rows=5, cols=3)
    table3.style = 'Table Grid'
    di_headers = ['Category', 'FY2025', 'FY2024']
    for i, h in enumerate(di_headers):
        cell = table3.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    di_data = [
        ['Women in Leadership', '38%', '34%'],
        ['Underrepresented Groups', '29%', '26%'],
        ['Gender Pay Gap', '< 2%', '3.1%'],
        ['Accessibility Compliance', '100%', '97%'],
    ]
    for r, row_data in enumerate(di_data, 1):
        for c, val in enumerate(row_data):
            table3.cell(r, c).text = val

    # ===================== PAGE 6 =====================
    doc.add_page_break()

    doc.add_heading('Sustainability & ESG Initiatives', level=1)
    doc.add_paragraph(
        'Environmental, Social, and Governance considerations are embedded in our '
        'strategic decision-making. In FY2025, we achieved carbon neutrality for our '
        'data center operations through a combination of renewable energy procurement '
        'and verified carbon offsets. Our total carbon footprint decreased by 31% '
        'compared to FY2023 baseline.'
    )

    doc.add_paragraph(
        'Our community investment program contributed $4.2 million to STEM education '
        'initiatives across 15 countries. The employee volunteer program logged over '
        '12,000 hours of community service. We maintained our AA rating from MSCI ESG '
        'and improved our CDP Climate Change score from B to A-.'
    )

    # ===================== PAGE 7 =====================
    doc.add_page_break()

    doc.add_heading('Risk Management & Governance', level=1)
    doc.add_paragraph(
        'Our enterprise risk management framework identified and assessed 47 strategic, '
        'operational, financial, and compliance risks during FY2025. The Board Risk '
        'Committee met quarterly to review the risk register and oversee mitigation '
        'strategies. Key risk categories include cybersecurity threats, regulatory '
        'compliance across jurisdictions, supply chain disruptions, and talent retention '
        'in competitive markets.'
    )

    doc.add_paragraph(
        'Cybersecurity investments totaled $34.8 million, including the deployment of '
        'an AI-powered threat detection platform that reduced mean time to detection '
        'from 72 hours to 4.3 hours. We successfully passed SOC 2 Type II, ISO 27001, '
        'and FedRAMP audits with zero critical findings.'
    )

    doc.add_heading('Governance Highlights', level=2)
    gov_items = [
        'Board composition: 11 directors, 9 independent, 4 women',
        'Audit Committee met 8 times with full attendance',
        'Compensation benchmarked at 50th percentile of peer group',
        'Whistleblower hotline: 12 reports received, all investigated and resolved',
    ]
    for item in gov_items:
        doc.add_paragraph(item, style='List Bullet')

    # ===================== PAGE 8 =====================
    doc.add_page_break()

    doc.add_heading('Strategic Outlook & Forward Guidance', level=1)
    doc.add_paragraph(
        'Looking ahead to FY2026, we expect revenue growth of 18-22%, driven by '
        'continued expansion of CloudSync Enterprise and the full integration of '
        'DataStream Analytics. We are targeting operating margins of 19-20% and plan '
        'to invest $140 million in R&D to accelerate our AI and cloud innovation pipeline.'
    )

    doc.add_paragraph(
        'Strategic priorities for the coming year include geographic expansion into '
        'Southeast Asia and the Middle East, the launch of our industry-specific '
        'vertical solutions for healthcare and financial services, and deepening our '
        'partner ecosystem with 50 new certified implementation partners.'
    )

    doc.add_paragraph(
        'Capital allocation will balance growth investment with shareholder returns. '
        'We plan to increase the quarterly dividend by 10% and continue our share '
        'repurchase program with a new $200 million authorization.'
    )

    # Last paragraph with [COMPANY] placeholder
    p_last = doc.add_paragraph()
    p_last.add_run('This report has been prepared and approved by the Board of Directors of ')
    run_c_last = p_last.add_run('[COMPANY]')
    run_c_last.bold = True
    p_last.add_run('. For questions or additional information, please contact the Office of '
                   'the Chief Financial Officer or the Investor Relations department.')

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
