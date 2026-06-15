"""
Reward Script: Create user variable 'CompanyName' and insert show-variable fields
Task ID: writer_tm_075
Domain: libreoffice_writer
Scoring:
  C1 (0.20) - Document variable 'CompanyName' = 'Nexus Technologies' defined in settings.xml
  C2 (0.30) - 3 DOCVARIABLE CompanyName field codes in document body
  C3 (0.15) - 1 DOCVARIABLE CompanyName field code in header
  C4 (0.15) - 1 DOCVARIABLE CompanyName field code in footer
  C5 (0.20) - Displayed text: '[COMPANY]' placeholders replaced with 'Nexus Technologies'
"""

import os
import zipfile
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_075'


def persist_app_state(domain: str):
    """Try to save any open LibreOffice document before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must exist and be a valid docx (zip)
    if not os.path.exists(file_path):
        print(f"CRITICAL: File not found: {file_path}")
        print("REWARD: 0.0")
        return 0.0

    try:
        zf = zipfile.ZipFile(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot open as ZIP: {e}")
        print("REWARD: 0.0")
        return 0.0

    # =========================================================================
    # Component 1: Document variable 'CompanyName' = 'Nexus Technologies' (0.20)
    # The task requires creating a user variable. In .docx format this is stored
    # as <w:docVar> in word/settings.xml.
    # =========================================================================
    try:
        settings_xml = zf.read('word/settings.xml').decode('utf-8')
        import re
        # Look for <w:docVar w:name="CompanyName" w:val="..."/>
        docvar_pattern = re.compile(
            r'<w:docVar\s+w:name="CompanyName"\s+w:val="([^"]*)"',
            re.IGNORECASE
        )
        match = docvar_pattern.search(settings_xml)
        if match:
            var_value = match.group(1)
            if var_value.strip() == 'Nexus Technologies':
                print(f"PASS: C1 — docVar CompanyName = '{var_value}' (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: C1 — docVar CompanyName has wrong value: '{var_value}'")
        else:
            print("FAIL: C1 — No docVar 'CompanyName' found in settings.xml")
    except Exception as e:
        print(f"ERROR: C1 — {e}")

    # =========================================================================
    # Component 2: 3 DOCVARIABLE CompanyName field codes in document body (0.30)
    # The task says: first paragraph, page 3, last paragraph.
    # These should be DOCVARIABLE field codes, not plain text.
    # =========================================================================
    try:
        from docx import Document
        from lxml import etree
        doc = Document(file_path)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        body_instrs = doc.element.body.findall('.//w:instrText', ns)
        body_docvar_count = 0
        for instr in body_instrs:
            if instr.text and 'DOCVARIABLE' in instr.text and 'CompanyName' in instr.text:
                body_docvar_count += 1

        if body_docvar_count >= 3:
            print(f"PASS: C2 — {body_docvar_count} DOCVARIABLE CompanyName fields in body (0.30 pts)")
            total_score += 0.30
        elif body_docvar_count >= 2:
            partial = 0.20
            print(f"PARTIAL: C2 — {body_docvar_count}/3 DOCVARIABLE fields in body ({partial} pts)")
            total_score += partial
        elif body_docvar_count >= 1:
            partial = 0.10
            print(f"PARTIAL: C2 — {body_docvar_count}/3 DOCVARIABLE fields in body ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: C2 — No DOCVARIABLE CompanyName fields found in body")
    except Exception as e:
        print(f"ERROR: C2 — {e}")

    # =========================================================================
    # Component 3: 1 DOCVARIABLE CompanyName field code in header (0.15)
    # =========================================================================
    try:
        header_docvar_count = 0
        for section in doc.sections:
            hdr_elem = section.header._element
            hdr_instrs = hdr_elem.findall('.//w:instrText', ns)
            for instr in hdr_instrs:
                if instr.text and 'DOCVARIABLE' in instr.text and 'CompanyName' in instr.text:
                    header_docvar_count += 1

        if header_docvar_count >= 1:
            print(f"PASS: C3 — {header_docvar_count} DOCVARIABLE CompanyName field(s) in header (0.15 pts)")
            total_score += 0.15
        else:
            print("FAIL: C3 — No DOCVARIABLE CompanyName field in header")
    except Exception as e:
        print(f"ERROR: C3 — {e}")

    # =========================================================================
    # Component 4: 1 DOCVARIABLE CompanyName field code in footer (0.15)
    # =========================================================================
    try:
        footer_docvar_count = 0
        for section in doc.sections:
            ftr_elem = section.footer._element
            ftr_instrs = ftr_elem.findall('.//w:instrText', ns)
            for instr in ftr_instrs:
                if instr.text and 'DOCVARIABLE' in instr.text and 'CompanyName' in instr.text:
                    footer_docvar_count += 1

        if footer_docvar_count >= 1:
            print(f"PASS: C4 — {footer_docvar_count} DOCVARIABLE CompanyName field(s) in footer (0.15 pts)")
            total_score += 0.15
        else:
            print("FAIL: C4 — No DOCVARIABLE CompanyName field in footer")
    except Exception as e:
        print(f"ERROR: C4 — {e}")

    # =========================================================================
    # Component 5: '[COMPANY]' placeholders replaced — text shows 'Nexus Technologies' (0.20)
    # Checks that the displayed text (including cached field results) contains
    # 'Nexus Technologies' and no remaining '[COMPANY]' placeholders.
    # =========================================================================
    try:
        all_text = []
        for para in doc.paragraphs:
            all_text.append(para.text)
        for section in doc.sections:
            for p in section.header.paragraphs:
                all_text.append(p.text)
            for p in section.footer.paragraphs:
                all_text.append(p.text)

        full_text = '\n'.join(all_text)
        has_placeholder = '[COMPANY]' in full_text
        has_nexus = 'Nexus Technologies' in full_text

        if has_nexus and not has_placeholder:
            print("PASS: C5 — 'Nexus Technologies' present, no '[COMPANY]' placeholders remain (0.20 pts)")
            total_score += 0.20
        elif has_nexus and has_placeholder:
            # Partial: some replaced, some not
            print("PARTIAL: C5 — 'Nexus Technologies' present but '[COMPANY]' placeholders still remain (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: C5 — has_nexus={has_nexus}, has_placeholder={has_placeholder}")
    except Exception as e:
        print(f"ERROR: C5 — {e}")

    zf.close()

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
