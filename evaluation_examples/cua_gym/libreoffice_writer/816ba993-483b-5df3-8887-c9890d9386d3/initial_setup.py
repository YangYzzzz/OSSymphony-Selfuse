"""
Initial Setup: Create a technical specification document without version history
Task ID: writer_tech_028
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_028'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('CloudSync API Technical Specification', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle / metadata ---
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run('Project: CloudSync Platform v3.0')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = meta2.add_run('Prepared by: Engineering Division, Nextera Technologies')
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- 1. Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This document defines the technical specification for the CloudSync API, '
        'a RESTful service enabling real-time data synchronization across distributed '
        'client applications. The API supports both push and pull synchronization models '
        'and is designed to handle up to 50,000 concurrent connections per node.'
    )
    doc.add_paragraph(
        'CloudSync API is the backbone of the CloudSync Platform v3.0, replacing the '
        'legacy SOAP-based integration layer that was deprecated in Q2 2025. This new '
        'architecture addresses performance bottlenecks identified during the v2.x '
        'load testing phase, particularly around conflict resolution latency.'
    )

    # --- 2. Scope ---
    doc.add_heading('2. Scope', level=1)
    doc.add_paragraph(
        'This specification covers the following components of the CloudSync API:'
    )
    doc.add_paragraph('Real-time bidirectional data synchronization engine', style='List Bullet')
    doc.add_paragraph('OAuth 2.0 authentication and role-based access control', style='List Bullet')
    doc.add_paragraph('Conflict resolution strategies (last-write-wins, merge, manual)', style='List Bullet')
    doc.add_paragraph('Rate limiting and throttling mechanisms', style='List Bullet')
    doc.add_paragraph('WebSocket and Server-Sent Events transport layers', style='List Bullet')
    doc.add_paragraph(
        'Out of scope: client-side SDK implementations, UI components, and third-party '
        'webhook integrations (covered in separate documents).'
    )

    # --- 3. System Requirements ---
    doc.add_heading('3. System Requirements', level=1)

    doc.add_heading('3.1 Hardware Requirements', level=2)
    doc.add_paragraph(
        'Each API node requires a minimum of 8 CPU cores (AMD EPYC 7543 or equivalent), '
        '32 GB RAM, and 500 GB NVMe SSD storage. The recommended configuration for '
        'production deployments is 16 cores, 64 GB RAM, and 1 TB storage with RAID-10.'
    )

    doc.add_heading('3.2 Software Requirements', level=2)
    doc.add_paragraph(
        'The API runtime environment requires Python 3.11+, PostgreSQL 15.x for persistent '
        'storage, Redis 7.x for session caching, and Kubernetes 1.28+ for container '
        'orchestration. TLS 1.3 is mandatory for all external communications.'
    )

    # --- 4. Architecture Overview ---
    doc.add_heading('4. Architecture Overview', level=1)
    doc.add_paragraph(
        'The CloudSync API follows a microservices architecture with the following core services:'
    )
    doc.add_paragraph('Sync Gateway: Routes and manages synchronization requests', style='List Number')
    doc.add_paragraph('Conflict Resolver: Handles merge conflicts using configurable strategies', style='List Number')
    doc.add_paragraph('Event Broker: Manages real-time event distribution via WebSocket/SSE', style='List Number')
    doc.add_paragraph('Auth Service: Handles OAuth 2.0 token lifecycle and RBAC', style='List Number')
    doc.add_paragraph('Data Store: Provides transactional persistence with PostgreSQL', style='List Number')
    doc.add_paragraph(
        'Inter-service communication uses gRPC with Protocol Buffers for internal calls, '
        'while external clients interact through the REST API gateway. All services are '
        'deployed as stateless containers managed by Kubernetes, with horizontal pod '
        'autoscaling configured based on CPU and memory thresholds.'
    )

    # --- 5. API Endpoints ---
    doc.add_heading('5. API Endpoints', level=1)
    doc.add_paragraph(
        'The primary API endpoints are organized under the /api/v3/ namespace. '
        'Authentication is required for all endpoints except /api/v3/health and '
        '/api/v3/status. Key endpoints include:'
    )
    doc.add_paragraph('POST /api/v3/sync/push - Submit local changes for synchronization', style='List Bullet')
    doc.add_paragraph('GET /api/v3/sync/pull - Retrieve pending changes from the server', style='List Bullet')
    doc.add_paragraph('POST /api/v3/auth/token - Obtain OAuth 2.0 access token', style='List Bullet')
    doc.add_paragraph('DELETE /api/v3/sync/conflicts/{id} - Resolve a specific conflict', style='List Bullet')
    doc.add_paragraph('GET /api/v3/ws/subscribe - Establish WebSocket connection for real-time events', style='List Bullet')

    # --- 6. Testing Strategy ---
    doc.add_heading('6. Testing Strategy', level=1)
    doc.add_paragraph(
        'Quality assurance for the CloudSync API involves four testing layers: unit tests '
        'targeting 90% code coverage, integration tests validating service interactions, '
        'load tests simulating peak traffic of 50,000 concurrent users, and chaos engineering '
        'experiments to verify fault tolerance under network partition scenarios.'
    )
    doc.add_paragraph(
        'All tests are executed in CI/CD pipelines using GitHub Actions. Performance '
        'regression tests run nightly against the staging environment, with automated '
        'alerts configured for latency increases exceeding 15% from the baseline p99.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
