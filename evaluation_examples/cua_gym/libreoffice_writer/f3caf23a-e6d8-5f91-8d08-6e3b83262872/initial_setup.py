"""
Initial Setup: Format a job description document with consistent styling
Task ID: writer_hr_028
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_028'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # All text uses Default Paragraph Style — no headings, no list styles, no formatting
    # Job title (plain text, NOT a heading)
    doc.add_paragraph('Product Manager')

    doc.add_paragraph('')  # blank line separator

    # Responsibilities section label (plain text)
    doc.add_paragraph('Responsibilities')

    # Bullet items as plain paragraphs (no List Bullet style)
    doc.add_paragraph('Define and prioritize the product roadmap based on market research, customer feedback, and business objectives')
    doc.add_paragraph('Collaborate with engineering, design, and marketing teams to deliver high-quality product releases on schedule')
    doc.add_paragraph('Conduct competitive analysis and identify market opportunities to drive product innovation')
    doc.add_paragraph('Gather and synthesize user requirements through interviews, surveys, and data analytics')
    doc.add_paragraph('Create detailed product specifications, user stories, and acceptance criteria for development teams')
    doc.add_paragraph('Monitor key product metrics and KPIs to evaluate feature performance and inform future decisions')

    doc.add_paragraph('')  # blank line separator

    # Qualifications section label (plain text)
    doc.add_paragraph('Qualifications')

    doc.add_paragraph('Bachelor\'s degree in Business, Computer Science, or a related field; MBA preferred')
    doc.add_paragraph('5+ years of experience in product management within a technology company')
    doc.add_paragraph('Strong analytical skills with proficiency in data-driven decision making')
    doc.add_paragraph('Excellent written and verbal communication skills with the ability to influence cross-functional stakeholders')
    doc.add_paragraph('Experience with Agile/Scrum methodologies and project management tools such as Jira or Asana')
    doc.add_paragraph('Demonstrated ability to manage multiple projects simultaneously and meet tight deadlines')

    doc.add_paragraph('')  # blank line separator

    # Benefits section label (plain text)
    doc.add_paragraph('Benefits')

    doc.add_paragraph('Competitive base salary ranging from $120,000 to $160,000 depending on experience')
    doc.add_paragraph('Annual performance bonus of up to 20% of base salary')
    doc.add_paragraph('Comprehensive health, dental, and vision insurance for employees and dependents')
    doc.add_paragraph('401(k) retirement plan with 6% company match')
    doc.add_paragraph('Flexible remote work policy with home office stipend of $1,500')
    doc.add_paragraph('Professional development budget of $3,000 per year for conferences and courses')

    doc.add_paragraph('')  # blank line separator

    # About Us section label (plain text)
    doc.add_paragraph('About Us')

    doc.add_paragraph('Apex Technologies is a leading SaaS company specializing in enterprise workflow automation solutions')
    doc.add_paragraph('Founded in 2015, we have grown to over 800 employees across offices in San Francisco, Austin, and London')
    doc.add_paragraph('Our platform serves more than 2,500 enterprise clients including Fortune 500 companies')
    doc.add_paragraph('We are committed to fostering an inclusive workplace where diverse perspectives drive innovation')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
