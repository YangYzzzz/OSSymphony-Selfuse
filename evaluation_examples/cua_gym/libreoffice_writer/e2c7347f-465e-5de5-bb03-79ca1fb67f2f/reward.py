"""
Reward Script: Donation Receipt in LibreOffice Writer
Task ID: writer_wf_058
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Org name "Hope Foundation" in 20pt bold
  Component 2 (0.10): Address below org name
  Component 3 (0.15): Title "Official Donation Receipt"
  Component 4 (0.25): Table with 7 required fields
  Component 5 (0.10): Tax-deductibility paragraph
  Component 6 (0.10): Signature line with name and title
  Component 7 (0.10): Italic "Thank You" at bottom
"""

import os
import re
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_058'


def persist_app_state(domain):
    """Save any unsaved changes in LibreOffice before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify donation receipt creation with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_text = '\n'.join(p.text for p in doc.paragraphs)

    # Component 1: Organization name "Hope Foundation" in 20pt bold (0.20 points)
    try:
        found_org = False
        for para in doc.paragraphs:
            for run in para.runs:
                if 'hope foundation' in run.text.lower():
                    is_bold = run.bold is True
                    size_ok = (run.font.size is not None and
                               abs(run.font.size.pt - 20.0) < 1.0)
                    if is_bold and size_ok:
                        found_org = True
                        print(f"PASS: Component 1 — 'Hope Foundation' found, bold={run.bold}, size={run.font.size.pt}pt (0.20 pts)")
                        total_score += 0.20
                        break
                    else:
                        print(f"PARTIAL: Component 1 — 'Hope Foundation' found but bold={run.bold}, size={run.font.size.pt if run.font.size else None}pt")
            if found_org:
                break
        if not found_org:
            print("FAIL: Component 1 — 'Hope Foundation' not found with 20pt bold")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Address below org name (0.10 points)
    try:
        # Check that there is address-like content (street, city, phone/email)
        found_address = False
        for para in doc.paragraphs:
            text_lower = para.text.lower()
            # Address should contain street-like info
            if any(kw in text_lower for kw in ['lane', 'street', 'drive', 'avenue', 'road', 'suite', 'blvd']):
                found_address = True
                print(f"PASS: Component 2 — Address found: '{para.text[:60]}...' (0.10 pts)")
                total_score += 0.10
                break
        if not found_address:
            print("FAIL: Component 2 — No address paragraph found below organization name")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Title "Official Donation Receipt" (0.15 points)
    try:
        found_title = False
        for para in doc.paragraphs:
            if 'official donation receipt' in para.text.lower():
                found_title = True
                # Check for bold or larger font
                has_emphasis = False
                for run in para.runs:
                    if run.bold is True or (run.font.size is not None and run.font.size.pt >= 14):
                        has_emphasis = True
                if has_emphasis:
                    print(f"PASS: Component 3 — Title 'Official Donation Receipt' found with emphasis (0.15 pts)")
                    total_score += 0.15
                else:
                    # Still give partial credit if title text exists
                    print(f"PARTIAL: Component 3 — Title text found but no bold/large formatting (0.10 pts)")
                    total_score += 0.10
                break
        if not found_title:
            print("FAIL: Component 3 — Title 'Official Donation Receipt' not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Table with 7 required fields (0.25 points)
    try:
        required_fields = [
            'receipt number',
            'date',
            'donor name',
            'donor address',
            'donation amount',
            'payment method',
            'purpose/fund'
        ]
        if len(doc.tables) == 0:
            print("FAIL: Component 4 — No tables found in document")
        else:
            # Find the table with the most matching fields
            best_match = 0
            best_table_idx = -1
            for ti, table in enumerate(doc.tables):
                field_texts = []
                for row in table.rows:
                    for cell in row.cells:
                        field_texts.append(cell.text.strip().lower())
                matches = sum(1 for f in required_fields if any(f in t for t in field_texts))
                if matches > best_match:
                    best_match = matches
                    best_table_idx = ti

            if best_match >= 7:
                print(f"PASS: Component 4 — Table with all 7 required fields found (0.25 pts)")
                total_score += 0.25
            elif best_match >= 5:
                pts = 0.15
                print(f"PARTIAL: Component 4 — Table has {best_match}/7 fields ({pts} pts)")
                total_score += pts
            elif best_match >= 3:
                pts = 0.10
                print(f"PARTIAL: Component 4 — Table has {best_match}/7 fields ({pts} pts)")
                total_score += pts
            elif best_match >= 1:
                pts = 0.05
                print(f"PARTIAL: Component 4 — Table has {best_match}/7 fields ({pts} pts)")
                total_score += pts
            else:
                print("FAIL: Component 4 — Table found but no matching fields")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Tax-deductibility paragraph (0.10 points)
    try:
        found_tax = False
        for para in doc.paragraphs:
            text_lower = para.text.lower()
            if 'tax' in text_lower and ('deducti' in text_lower or 'exempt' in text_lower or '501' in text_lower):
                found_tax = True
                print(f"PASS: Component 5 — Tax-deductibility statement found (0.10 pts)")
                total_score += 0.10
                break
        if not found_tax:
            print("FAIL: Component 5 — No tax-deductibility statement found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Signature line with name and title (0.10 points)
    try:
        # Look for a line of underscores (signature line) followed by name/title
        found_sig_line = False
        found_name_title = False
        for i, para in enumerate(doc.paragraphs):
            if '____' in para.text:
                found_sig_line = True
                # Check subsequent paragraphs for name and title
                remaining = [doc.paragraphs[j].text for j in range(i + 1, min(i + 4, len(doc.paragraphs)))]
                remaining_lower = ' '.join(remaining).lower()
                if any(kw in remaining_lower for kw in ['director', 'president', 'manager', 'officer', 'executive', 'treasurer', 'secretary']):
                    found_name_title = True
                break

        if found_sig_line and found_name_title:
            print(f"PASS: Component 6 — Signature line with name and title found (0.10 pts)")
            total_score += 0.10
        elif found_sig_line:
            print(f"PARTIAL: Component 6 — Signature line found but no name/title below (0.05 pts)")
            total_score += 0.05
        else:
            print("FAIL: Component 6 — No signature line found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Italic "Thank You" message at bottom (0.10 points)
    try:
        # Check last few non-empty paragraphs for italic "thank you"
        non_empty = [p for p in doc.paragraphs if p.text.strip()]
        found_thankyou = False
        if non_empty:
            # Check last 3 non-empty paragraphs
            for para in non_empty[-3:]:
                if 'thank you' in para.text.lower() or 'thank' in para.text.lower():
                    # Check if italic
                    has_italic = any(run.italic is True for run in para.runs if run.text.strip())
                    if has_italic:
                        print(f"PASS: Component 7 — Italic 'Thank You' message found at bottom (0.10 pts)")
                        total_score += 0.10
                        found_thankyou = True
                        break
                    else:
                        print(f"PARTIAL: Component 7 — 'Thank You' found but not italic (0.05 pts)")
                        total_score += 0.05
                        found_thankyou = True
                        break
        if not found_thankyou:
            print("FAIL: Component 7 — No 'Thank You' message found at bottom")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
