"""
Initial Setup: HR Employee Roster - Tab Stop Alignment Task
Task ID: osworld_writer_tabstop_split_line_005
Domain: libreoffice_writer

Creates a document with an HR employee roster where each line has
employee name and department separated by spaces (no tab stops configured).
The agent's task is to apply left tab at 0 cm and right tab at 16 cm to
all roster lines.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_tabstop_split_line_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("Acme Corp — Employee Roster")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Subtitle / date
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = sub_para.add_run("Q1 2025 | Human Resources Department")
    sub_run.font.size = Pt(11)
    sub_run.italic = True

    # Blank line
    doc.add_paragraph()

    # Section header
    section_header = doc.add_paragraph()
    sec_run = section_header.add_run("Current Staff Listing")
    sec_run.bold = True
    sec_run.font.size = Pt(13)
    sec_run.underline = True

    # Blank line before roster
    doc.add_paragraph()

    # Roster data — 8 lines with name and department separated by spaces
    # No tab stops are configured; names and departments just separated by spaces
    roster = [
        ("Sarah Chen", "Engineering"),
        ("Marcus Johnson", "Marketing"),
        ("Priya Patel", "Finance"),
        ("David Nguyen", "Operations"),
        ("Olivia Martinez", "Human Resources"),
        ("James O'Brien", "Legal"),
        ("Aisha Kamara", "Product Design"),
        ("Robert Schultz", "Customer Success"),
    ]

    for name, dept in roster:
        para = doc.add_paragraph()
        # Use spaces to visually separate — NO tab characters, NO tab stops
        para.add_run(f"{name}    {dept}")

    # Blank line after roster
    doc.add_paragraph()

    # Footer note
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run("Confidential — For internal use only. Do not distribute.")
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
