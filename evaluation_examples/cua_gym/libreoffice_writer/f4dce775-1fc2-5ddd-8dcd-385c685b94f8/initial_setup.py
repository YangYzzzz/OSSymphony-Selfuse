"""
Initial Setup: Technical manual with specification table using inconsistent spacing separators
Task ID: osworld_writer_tabstop_split_line_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_tabstop_split_line_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    # --- Document Title ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("IndustrialEdge Controller IEC-7200")
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run("Technical Reference Manual")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)

    doc.add_paragraph()  # blank line

    # --- Section 1: Overview ---
    heading1 = doc.add_paragraph()
    h1_run = heading1.add_run("1. Product Overview")
    h1_run.bold = True
    h1_run.font.size = Pt(13)

    overview_text = (
        "The IEC-7200 is a DIN-rail mounted programmable logic controller designed for "
        "demanding industrial automation environments. Operating reliably across a wide "
        "temperature range, it supports both IEC 61131-3 structured text and ladder "
        "diagram programming paradigms. The onboard Ethernet interface enables seamless "
        "integration with SCADA systems and cloud-based monitoring platforms."
    )
    doc.add_paragraph(overview_text)

    doc.add_paragraph()

    # --- Section 2: Technical Specifications ---
    spec_heading = doc.add_paragraph()
    spec_run = spec_heading.add_run("2. Technical Specifications")
    spec_run.bold = True
    spec_run.font.size = Pt(13)

    intro_para = doc.add_paragraph(
        "The following table lists the key electrical and mechanical parameters of "
        "the IEC-7200 controller. All measurements are taken under nominal operating "
        "conditions unless otherwise stated."
    )

    doc.add_paragraph()

    # 18 specification lines using INCONSISTENT spacing separators (no tab stops, no \t)
    # Parameter names and values separated by multiple spaces (inconsistent)
    spec_lines = [
        ("Supply Voltage",           "24 VDC ± 15%"),
        ("Power Consumption",        "18 W (max)"),
        ("CPU Architecture",         "ARM Cortex-A9 @ 800 MHz"),
        ("RAM",                      "512 MB DDR3"),
        ("Flash Storage",            "4 GB eMMC"),
        ("Operating Temperature",    "-20°C to +60°C"),
        ("Storage Temperature",      "-40°C to +85°C"),
        ("Relative Humidity",        "5% to 95% (non-condensing)"),
        ("Protection Rating",        "IP20"),
        ("Digital Inputs",           "16 x 24 VDC, 3 ms filter"),
        ("Digital Outputs",          "8 x relay, 2 A / 250 VAC"),
        ("Analog Inputs",            "4 x 0–10 V / 4–20 mA, 12-bit"),
        ("Analog Outputs",           "2 x 0–10 V, 12-bit"),
        ("Communication Ports",      "2 x RS-485, 1 x Ethernet 100Base-T"),
        ("Programming Standard",     "IEC 61131-3"),
        ("Dimensions (W × H × D)",   "75 mm × 120 mm × 58 mm"),
        ("Weight",                   "420 g"),
        ("Certifications",           "CE, UL 508, RoHS"),
    ]

    # Use inconsistent spacing - different numbers of spaces for each line
    # to simulate real-world messy formatting
    import random
    random.seed(42)
    space_patterns = [
        "    ",     # 4 spaces
        "      ",   # 6 spaces
        "   ",      # 3 spaces
        "     ",    # 5 spaces
        "        ", # 8 spaces
        "  ",       # 2 spaces
        "       ",  # 7 spaces
        "          ", # 10 spaces
        "    ",
        "      ",
        "   ",
        "     ",
        "        ",
        "  ",
        "       ",
        "          ",
        "    ",
        "      ",
    ]

    for i, (param, value) in enumerate(spec_lines):
        p = doc.add_paragraph()
        # No tab stops, just spaces as separator - inconsistent spacing
        spaces = space_patterns[i % len(space_patterns)]
        run = p.add_run(f"{param}{spaces}{value}")
        run.font.size = Pt(11)

    doc.add_paragraph()

    # --- Section 3: Installation ---
    inst_heading = doc.add_paragraph()
    inst_run = inst_heading.add_run("3. Installation Guidelines")
    inst_run.bold = True
    inst_run.font.size = Pt(13)

    inst_text = (
        "Mount the IEC-7200 on a standard 35 mm DIN rail. Ensure a minimum clearance "
        "of 50 mm above and below the unit for adequate ventilation. Connect the 24 VDC "
        "supply to terminals L+ and M. Do not exceed the specified supply voltage range "
        "as this may void the warranty and damage the device."
    )
    doc.add_paragraph(inst_text)

    doc.add_paragraph()

    # --- Section 4: Wiring ---
    wiring_heading = doc.add_paragraph()
    wiring_run = wiring_heading.add_run("4. Wiring Diagrams and I/O Configuration")
    wiring_run.bold = True
    wiring_run.font.size = Pt(13)

    wiring_text = (
        "All field wiring connections use spring-cage terminals rated for conductors "
        "from 0.5 mm² to 2.5 mm². Tighten terminal screws to the torque specified on "
        "the terminal cover label. Shielded cables are recommended for analog signals "
        "longer than 3 meters. Ground the cable shield at one end only to prevent "
        "ground loop interference."
    )
    doc.add_paragraph(wiring_text)

    doc.add_paragraph()

    # --- Section 5: Maintenance ---
    maint_heading = doc.add_paragraph()
    maint_run = maint_heading.add_run("5. Maintenance and Troubleshooting")
    maint_run.bold = True
    maint_run.font.size = Pt(13)

    maint_text = (
        "The IEC-7200 requires no periodic maintenance under normal operating conditions. "
        "Clean the device exterior with a dry cloth only; do not use solvents. The onboard "
        "real-time clock battery (CR2032) should be replaced every five years or when the "
        "low-battery indicator activates. For firmware updates, connect via the Ethernet "
        "port and use the IEC Manager configuration tool version 3.4 or later."
    )
    doc.add_paragraph(maint_text)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
