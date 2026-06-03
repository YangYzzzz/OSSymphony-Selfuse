"""
Reward Script: Apply right-aligned tab stops at 16cm to 18 specification lines in technical manual
Task ID: osworld_writer_tabstop_split_line_009
Domain: libreoffice_writer
Scoring:
  Component 1: All 18 spec lines have a RIGHT tab stop at ~16cm (0.5 points)
  Component 2: All 18 spec lines use a tab character as the separator (0.5 points)
"""

import os
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_tabstop_split_line_009'

# Specification lines are paragraphs 9 through 26 (indices, 0-based)
SPEC_PARA_INDICES = list(range(9, 27))  # 18 lines
NUM_SPEC_LINES = 18

# Expected right tab stop position in EMU (~16.00 cm).
# Actual observed value in golden: 5760085 EMU ≈ 16.0002 cm.
# Allow ±1% tolerance around exactly 16 cm (914400 EMU/inch * 16/2.54 inches)
RIGHT_TAB_EXPECTED_CM = 16.0
EMU_PER_CM = 914400 / 2.54  # 360000 EMU per cm
RIGHT_TAB_EXPECTED_EMU = RIGHT_TAB_EXPECTED_CM * EMU_PER_CM  # 5760000 EMU
RIGHT_TAB_TOLERANCE_EMU = RIGHT_TAB_EXPECTED_EMU * 0.01  # 1% tolerance (~57600 EMU, ~0.16 cm)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: document must have at least 27 paragraphs
    if len(doc.paragraphs) < 27:
        print(f"CRITICAL: Document has only {len(doc.paragraphs)} paragraphs, expected at least 27")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: All 18 specification lines have a RIGHT tab stop at ~16cm
    # This FAILS on initial (no tab stops at all) → PASSES on golden
    # -------------------------------------------------------------------------
    try:
        lines_with_right_tab = 0
        lines_missing_right_tab = []

        for idx in SPEC_PARA_INDICES:
            para = doc.paragraphs[idx]
            has_right_tab_at_16cm = False
            for ts in para.paragraph_format.tab_stops:
                # Skip CLEAR and default LEFT@0 stops
                if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
                    continue
                if ts.alignment == WD_TAB_ALIGNMENT.LEFT and ts.position == 0:
                    continue
                # Check for RIGHT alignment at ~16cm
                if ts.alignment == WD_TAB_ALIGNMENT.RIGHT:
                    pos_cm = ts.position / EMU_PER_CM
                    if abs(ts.position - RIGHT_TAB_EXPECTED_EMU) <= RIGHT_TAB_TOLERANCE_EMU:
                        has_right_tab_at_16cm = True
                        break

            if has_right_tab_at_16cm:
                lines_with_right_tab += 1
            else:
                lines_missing_right_tab.append(idx)

        if lines_with_right_tab == NUM_SPEC_LINES:
            print(f"PASS: Component 1 — All {NUM_SPEC_LINES} spec lines have RIGHT tab stop at ~16cm (0.5 pts)")
            total_score += 0.5
        elif lines_with_right_tab > 0:
            # Partial: some lines have the right tab stop
            partial = round((lines_with_right_tab / NUM_SPEC_LINES) * 0.5, 4)
            print(f"PARTIAL: Component 1 — {lines_with_right_tab}/{NUM_SPEC_LINES} lines have RIGHT tab @16cm ({partial} pts)")
            print(f"  Missing on para indices: {lines_missing_right_tab[:5]}{'...' if len(lines_missing_right_tab) > 5 else ''}")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No spec lines have RIGHT tab stop at ~16cm (found: 0/{NUM_SPEC_LINES})")

    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: All 18 specification lines use a tab character as separator
    # (spacing separators replaced with \t)
    # This FAILS on initial (uses spaces) → PASSES on golden (uses \t)
    # -------------------------------------------------------------------------
    try:
        lines_with_tab_char = 0
        lines_missing_tab_char = []

        for idx in SPEC_PARA_INDICES:
            para = doc.paragraphs[idx]
            full_text = para.text
            if '\t' in full_text:
                lines_with_tab_char += 1
            else:
                lines_missing_tab_char.append(idx)

        if lines_with_tab_char == NUM_SPEC_LINES:
            print(f"PASS: Component 2 — All {NUM_SPEC_LINES} spec lines use tab character as separator (0.5 pts)")
            total_score += 0.5
        elif lines_with_tab_char > 0:
            partial = round((lines_with_tab_char / NUM_SPEC_LINES) * 0.5, 4)
            print(f"PARTIAL: Component 2 — {lines_with_tab_char}/{NUM_SPEC_LINES} lines use tab character ({partial} pts)")
            print(f"  Missing on para indices: {lines_missing_tab_char[:5]}{'...' if len(lines_missing_tab_char) > 5 else ''}")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No spec lines use tab character (still using spacing separators)")

    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM environment
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
