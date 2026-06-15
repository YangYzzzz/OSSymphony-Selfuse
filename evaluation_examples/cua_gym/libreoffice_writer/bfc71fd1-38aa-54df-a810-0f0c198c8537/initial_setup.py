"""
Initial Setup: Certificate of Achievement document without page borders
Task ID: writer_rd_026
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_026'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup: A4 with 2.54 cm margins ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # --- Add some vertical spacing at top ---
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(60)
    spacer.paragraph_format.space_after = Pt(0)

    # --- Title ---
    title = doc.add_heading('Certificate of Achievement', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        run.font.size = Pt(28)

    # --- Decorative line ---
    line_para = doc.add_paragraph()
    line_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    line_para.paragraph_format.space_before = Pt(6)
    line_para.paragraph_format.space_after = Pt(24)
    line_run = line_para.add_run('\u2014' * 30)
    line_run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    line_run.font.size = Pt(12)

    # --- Presented to ---
    presented = doc.add_paragraph()
    presented.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    presented.paragraph_format.space_before = Pt(12)
    presented.paragraph_format.space_after = Pt(6)
    run_p = presented.add_run('This certificate is proudly presented to')
    run_p.font.size = Pt(14)
    run_p.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # --- Recipient Name ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_para.paragraph_format.space_before = Pt(12)
    name_para.paragraph_format.space_after = Pt(12)
    name_run = name_para.add_run('Alexandra Mikhailova')
    name_run.bold = True
    name_run.font.size = Pt(26)
    name_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # --- Description ---
    desc_para = doc.add_paragraph()
    desc_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    desc_para.paragraph_format.space_before = Pt(12)
    desc_para.paragraph_format.space_after = Pt(6)
    desc_run = desc_para.add_run(
        'In recognition of outstanding contributions to the International '
        'Data Science Research Initiative and exceptional leadership in '
        'cross-functional team collaboration during the 2025 fiscal year.'
    )
    desc_run.font.size = Pt(12)
    desc_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.space_before = Pt(24)
    date_para.paragraph_format.space_after = Pt(36)
    date_run = date_para.add_run('Awarded on March 15, 2025')
    date_run.font.size = Pt(11)
    date_run.font.italic = True
    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Signature Lines ---
    # Left signature
    sig1 = doc.add_paragraph()
    sig1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sig1.paragraph_format.space_before = Pt(40)
    sig1.paragraph_format.space_after = Pt(2)
    sig1_run = sig1.add_run('_________________________          _________________________')
    sig1_run.font.size = Pt(11)

    sig_names = doc.add_paragraph()
    sig_names.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sig_names.paragraph_format.space_before = Pt(4)
    sig_names.paragraph_format.space_after = Pt(2)
    sn_run = sig_names.add_run('Dr. Richard Thornton                    Prof. Elena Vasquez')
    sn_run.font.size = Pt(10)
    sn_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    sig_titles = doc.add_paragraph()
    sig_titles.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sig_titles.paragraph_format.space_before = Pt(0)
    sig_titles.paragraph_format.space_after = Pt(6)
    st_run = sig_titles.add_run('Director of Research                      Department Chair')
    st_run.font.size = Pt(9)
    st_run.font.italic = True
    st_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Organization name at bottom ---
    org = doc.add_paragraph()
    org.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    org.paragraph_format.space_before = Pt(30)
    org_run = org.add_run('Global Institute of Advanced Technology')
    org_run.font.size = Pt(11)
    org_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    org_run.font.small_caps = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
