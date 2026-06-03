"""
FINAL REWARD SCRIPT - SUCCESS
Task: My cursor is blinking where I want a table, and it has to be completely blank with exactly 7 columns and 5 rows (no header row). How do I drop that in quickly in LibreOffice Writer?
Generated: 2025-09-10 13:38:09
Status: success
Model: azure-o3
Total Steps: 2
"""

from docx import Document
import os

def verify_table_task(file_path: str) -> float:
    """Verify the LibreOffice Writer task:
    The document must contain a completely blank table with exactly
    7 columns and 5 rows (no header row).  Progressive scoring is
    applied based on how many requirements are met.
    Returns a float between 0.0-1.0.
    """

    print(f"Verifying document at: {file_path}")

    total_score = 0.0          # progressive score accumulator
    max_score   = 1.0          # cap to 1.0

    # ------------------------------------------------------------------
    # 1.  File existence & loading (NO POINTS – prerequisite only)
    # ------------------------------------------------------------------
    if not os.path.exists(file_path):
        print("✗ File does not exist.")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Unable to load DOCX file: {e}")
        return 0.0

    # ------------------------------------------------------------------
    # 2.  Check for presence of at least one table (0.2 pts)
    # ------------------------------------------------------------------
    tables = doc.tables
    if not tables:
        print("✗ No tables found in the document.")
        return 0.0  # cannot get any further points without a table

    print(f"✓ Found {len(tables)} table(s) in the document.")
    total_score += 0.2  # earned for having a table at all

    # ------------------------------------------------------------------
    # 3.  Find a table with exactly 5 rows and 7 columns (0.5 pts)
    #     • 0.25 for correct row count
    #     • 0.25 for correct column count
    # ------------------------------------------------------------------
    matching_table = None
    for t_idx, tbl in enumerate(tables, start=1):
        rows = len(tbl.rows)
        cols = len(tbl.columns) if tbl.rows else 0
        print(f"  Table {t_idx}: {rows} rows, {cols} columns")

        if rows == 5 and cols == 7:
            matching_table = tbl
            break

    if matching_table is None:
        print("✗ No table with exactly 5 rows and 7 columns found.")
    else:
        print("✓ Found a table with 5 rows and 7 columns.")
        total_score += 0.25  # correct row count
        total_score += 0.25  # correct column count

        # --------------------------------------------------------------
        # 4.  Verify that EVERY cell in this table is completely blank
        #     (0.3 pts)
        # --------------------------------------------------------------
        all_blank = True
        for r_idx, row in enumerate(matching_table.rows, start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                # Concatenate text from all paragraphs inside the cell
                cell_text = "".join(p.text for p in cell.paragraphs).strip()
                if cell_text:  # any non-empty cell violates requirement
                    print(
                        f"✗ Cell at row {r_idx}, col {c_idx} not blank: '{cell_text}'"
                    )
                    all_blank = False
                    break
            if not all_blank:
                break

        if all_blank:
            print("✓ All cells in the table are blank.")
            total_score += 0.3
        else:
            print("✗ Not all cells are blank in the table.")

    # ------------------------------------------------------------------
    # 5.  Final score (capped at 1.0)
    # ------------------------------------------------------------------
    final_score = min(total_score, max_score)
    print(f"Total score: {final_score}")
    return final_score


if __name__ == "__main__":
    # Path provided in the task context
    file_path = (
        "/home/user/"
        "my_cursor_is_blinking_where_i_want_a_table_and_it_has_to_be_completely_blank_with_exactly_7_columns_.docx"
    )

    reward = verify_table_task(file_path)
    print(f"REWARD: {reward}")
