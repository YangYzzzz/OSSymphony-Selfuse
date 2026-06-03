"""
Reward Script: Safety Data Sheet (SDS) for CleanMax Pro
Task ID: writer_wf_040
Domain: libreoffice_writer
Scoring:
  Component 1: Red DANGER header (0.15)
  Component 2: 6 Heading 1 sections (0.25)
  Component 3: Product Identification table (0.15)
  Component 4: Composition table with 3 ingredients (0.15)
  Component 5: First Aid sub-sections (0.15)
  Component 6: Storage and Disposal content (0.15)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_040'


def persist_app_state(domain: str):
    """Save any unsaved LibreOffice state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.shared import RGBColor
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather document data
    paragraphs = doc.paragraphs
    tables = doc.tables

    # If document is essentially empty, fast exit
    if len(paragraphs) == 0 and len(tables) == 0:
        print("FAIL: Document is empty — no paragraphs or tables found")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Red "DANGER" header (0.15 points)
    # Task requires a red 'DANGER' header at the top of the document
    try:
        danger_found = False
        for para in paragraphs:
            text = para.text.strip().upper()
            if 'DANGER' in text:
                for run in para.runs:
                    if 'DANGER' in run.text.upper():
                        rgb = run.font.color.rgb if run.font.color and run.font.color.rgb else None
                        if rgb is not None:
                            # Check if the color is red-ish (R high, G and B low)
                            r_val = int(str(rgb)[:2], 16)
                            g_val = int(str(rgb)[2:4], 16)
                            b_val = int(str(rgb)[4:6], 16)
                            if r_val >= 200 and g_val < 80 and b_val < 80:
                                danger_found = True
                                break
                if danger_found:
                    break
        if danger_found:
            print(f"PASS: Component 1 — Red 'DANGER' header found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Red 'DANGER' header not found or not red")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 6 Heading 1 sections (0.25 points)
    # Task requires 6 sections with Heading 1 style covering specific topics
    try:
        heading1_texts = []
        for para in paragraphs:
            if para.style and para.style.name == 'Heading 1':
                heading1_texts.append(para.text.strip().lower())

        required_sections = [
            'product identification',
            'hazard identification',
            'composition',
            'first aid',
            'storage',
            'disposal',
        ]

        matched_sections = 0
        for req in required_sections:
            for h1 in heading1_texts:
                if req in h1:
                    matched_sections += 1
                    break

        if matched_sections == 6:
            print(f"PASS: Component 2 — All 6 Heading 1 sections found ({heading1_texts}) (0.25 pts)")
            total_score += 0.25
        elif matched_sections >= 4:
            partial = round(0.25 * (matched_sections / 6), 2)
            print(f"PARTIAL: Component 2 — {matched_sections}/6 sections found (matched: {heading1_texts}) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {matched_sections}/6 required Heading 1 sections. Found: {heading1_texts}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Product Identification table (0.15 points)
    # Must have a table with Product Name, Manufacturer, Emergency Contact
    try:
        prod_table_found = False
        for table in tables:
            cell_texts = []
            for row in table.rows:
                for cell in row.cells:
                    cell_texts.append(cell.text.strip().lower())

            has_product_name = any('product name' in t for t in cell_texts)
            has_manufacturer = any('manufacturer' in t for t in cell_texts)
            has_emergency = any('emergency' in t for t in cell_texts)
            has_cleanmax = any('cleanmax' in t for t in cell_texts)

            if has_product_name and has_manufacturer and has_emergency and has_cleanmax:
                prod_table_found = True
                break

        if prod_table_found:
            print(f"PASS: Component 3 — Product Identification table found with required fields (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 — Product Identification table missing or incomplete")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Composition table with 3 ingredients (0.15 points)
    # Must have table with Component, CAS Number, Percentage columns and 3 data rows
    try:
        comp_table_found = False
        for table in tables:
            # Check if this is the composition table
            header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells] if len(table.rows) > 0 else []
            has_component = any('component' in t for t in header_texts)
            has_cas = any('cas' in t for t in header_texts)
            has_percentage = any('percent' in t or '%' in t for t in header_texts)

            if has_component and has_cas and has_percentage:
                # Count data rows (excluding header)
                data_rows = len(table.rows) - 1
                if data_rows >= 3:
                    # Verify data rows have content
                    filled_rows = 0
                    for row in list(table.rows)[1:]:
                        row_text = ' '.join(cell.text.strip() for cell in row.cells)
                        if len(row_text) > 5:  # non-trivial content
                            filled_rows += 1
                    if filled_rows >= 3:
                        comp_table_found = True
                        print(f"  Composition table: {data_rows} data rows, {filled_rows} filled")
                        break

        if comp_table_found:
            print(f"PASS: Component 4 — Composition table with 3+ ingredients found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 4 — Composition table missing or has fewer than 3 ingredients")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: First Aid sub-sections (0.15 points)
    # Must have sub-sections for inhalation, skin, eyes, ingestion
    try:
        heading2_texts = []
        for para in paragraphs:
            if para.style and para.style.name == 'Heading 2':
                heading2_texts.append(para.text.strip().lower())

        required_subsections = ['inhalation', 'skin', 'eye', 'ingestion']
        matched_sub = 0
        for req in required_subsections:
            for h2 in heading2_texts:
                if req in h2:
                    matched_sub += 1
                    break

        if matched_sub >= 4:
            print(f"PASS: Component 5 — All 4 First Aid sub-sections found ({heading2_texts}) (0.15 pts)")
            total_score += 0.15
        elif matched_sub >= 2:
            partial = round(0.15 * (matched_sub / 4), 2)
            print(f"PARTIAL: Component 5 — {matched_sub}/4 sub-sections found ({heading2_texts}) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Only {matched_sub}/4 First Aid sub-sections found. Found headings: {heading2_texts}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Storage and Disposal content (0.15 points)
    # Must have substantive text under Storage/Handling and Disposal sections
    try:
        full_text = ' '.join(p.text.strip().lower() for p in paragraphs)
        has_storage_content = 'storage' in full_text and ('ventilat' in full_text or 'cool' in full_text or 'dry' in full_text)
        has_disposal_content = 'disposal' in full_text and ('regulat' in full_text or 'waste' in full_text or 'dispos' in full_text)

        if has_storage_content and has_disposal_content:
            print(f"PASS: Component 6 — Storage and Disposal content present (0.15 pts)")
            total_score += 0.15
        elif has_storage_content or has_disposal_content:
            print(f"PARTIAL: Component 6 — Only one of storage/disposal has content (0.075 pts)")
            total_score += 0.075
        else:
            print(f"FAIL: Component 6 — Storage and Disposal content missing or insufficient")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
