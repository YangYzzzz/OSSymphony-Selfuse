"""
Initial Setup: Sales Analysis document with table - Achievement % column empty
Task ID: writer_af_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_af_041'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("Quarterly Sales Analysis Report", level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        "The following table summarizes the actual versus target sales performance "
        "across all regions for Q1 2025. The Achievement % column needs to be filled in "
        "to show each region-product combination's performance against targets."
    )
    intro.paragraph_format.space_after = Pt(12)

    # Data: Region, Product, Actual Sales, Target Sales
    data_rows = [
        ("Northeast", "Electronics",  187500, 200000),
        ("Southeast", "Furniture",     92300, 110000),
        ("Midwest",   "Office Supplies", 45800,  42000),
        ("West Coast","Electronics",  215600, 220000),
        ("Northeast", "Furniture",    134200, 150000),
        ("Southeast", "Office Supplies", 67900,  65000),
        ("Midwest",   "Electronics",  156400, 180000),
        ("West Coast","Furniture",    108700, 100000),
    ]

    headers = ["Region", "Product", "Actual Sales", "Target Sales", "Achievement %"]

    # Create table: 1 header row + 8 data rows = 9 rows, 5 columns
    table = doc.add_table(rows=1 + len(data_rows), cols=5)
    table.style = "Table Grid"

    # Header row
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Light blue header background
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '4472C4',
        })
        shading.append(shading_elem)
        # White text for header
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for row_idx, (region, product, actual, target) in enumerate(data_rows, start=1):
        # Column A: Region
        cell_a = table.cell(row_idx, 0)
        cell_a.text = ""
        r = cell_a.paragraphs[0].add_run(region)
        r.font.size = Pt(10)
        r.font.name = "Calibri"

        # Column B: Product
        cell_b = table.cell(row_idx, 1)
        cell_b.text = ""
        r = cell_b.paragraphs[0].add_run(product)
        r.font.size = Pt(10)
        r.font.name = "Calibri"

        # Column C: Actual Sales
        cell_c = table.cell(row_idx, 2)
        cell_c.text = ""
        r = cell_c.paragraphs[0].add_run(f"${actual:,.0f}")
        r.font.size = Pt(10)
        r.font.name = "Calibri"
        cell_c.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Column D: Target Sales
        cell_d = table.cell(row_idx, 3)
        cell_d.text = ""
        r = cell_d.paragraphs[0].add_run(f"${target:,.0f}")
        r.font.size = Pt(10)
        r.font.name = "Calibri"
        cell_d.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Column E: Achievement % — LEFT EMPTY (task requires filling this)
        cell_e = table.cell(row_idx, 4)
        cell_e.text = ""
        cell_e.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.3)
        row.cells[1].width = Inches(1.5)
        row.cells[2].width = Inches(1.3)
        row.cells[3].width = Inches(1.3)
        row.cells[4].width = Inches(1.3)

    # Footer note
    doc.add_paragraph("")  # spacer
    footer_note = doc.add_paragraph(
        "Note: Achievement % should be calculated as (Actual Sales / Target Sales) x 100, "
        "displayed with one decimal place."
    )
    footer_note.runs[0].font.size = Pt(9)
    footer_note.runs[0].italic = True

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
