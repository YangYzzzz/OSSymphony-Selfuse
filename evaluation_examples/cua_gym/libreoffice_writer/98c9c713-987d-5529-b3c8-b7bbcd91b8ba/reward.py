"""
Reward Script: Add percentage formulas to column E of a Writer table
Task ID: writer_af_041
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): All 8 column E cells (rows 1-8) are non-empty
  Component 2 (0.4): Values match expected (Actual/Target)*100 rounded to 1 decimal with % sign
  Component 3 (0.2): Format consistency - all values have X.X% pattern
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_af_041'


def parse_currency(text):
    """Parse currency string like '$187,500' to a number."""
    cleaned = text.replace('$', '').replace(',', '').strip()
    try:
        return float(cleaned)
    except (ValueError, TypeError):
        return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Precondition: table must have 9 rows and 5 columns
    if len(table.rows) < 9 or len(table.columns) < 5:
        print(f"CRITICAL: Table dimensions wrong: {len(table.rows)} rows x {len(table.columns)} cols, expected 9x5")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All 8 column E cells (rows 1-8) are non-empty (0.4 points)
    try:
        non_empty_count = 0
        for ri in range(1, 9):
            cell_text = table.cell(ri, 4).text.strip()
            if cell_text:
                non_empty_count += 1

        if non_empty_count == 8:
            print(f"PASS: Component 1 — All 8 column E cells are filled ({non_empty_count}/8) (0.4 pts)")
            total_score += 0.4
        elif non_empty_count > 0:
            partial = round(0.4 * (non_empty_count / 8), 2)
            print(f"PARTIAL: Component 1 — {non_empty_count}/8 column E cells filled ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — All column E cells are empty (0/8)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Values match expected calculations (0.4 points)
    # Expected: (Actual Sales / Target Sales) * 100, rounded to 1 decimal, with % sign
    try:
        correct_count = 0
        for ri in range(1, 9):
            actual_text = table.cell(ri, 2).text.strip()  # Column C: Actual Sales
            target_text = table.cell(ri, 3).text.strip()  # Column D: Target Sales
            result_text = table.cell(ri, 4).text.strip()  # Column E: Achievement %

            actual_val = parse_currency(actual_text)
            target_val = parse_currency(target_text)

            if actual_val is None or target_val is None or target_val == 0:
                print(f"  Row {ri}: Cannot parse sales values: C={actual_text}, D={target_text}")
                continue

            expected_pct = round((actual_val / target_val) * 100, 1)
            expected_str = f"{expected_pct:.1f}%"

            # Also accept without leading/trailing spaces
            result_clean = result_text.strip()

            # Try to parse the result value
            pct_match = re.match(r'^([\d.]+)\s*%$', result_clean)
            if pct_match:
                result_val = float(pct_match.group(1))
                # Allow small tolerance for rounding differences
                if abs(result_val - expected_pct) < 0.15:
                    correct_count += 1
                    print(f"  Row {ri}: MATCH — expected {expected_str}, found {result_clean}")
                else:
                    print(f"  Row {ri}: MISMATCH — expected {expected_str}, found {result_clean}")
            else:
                print(f"  Row {ri}: BAD FORMAT — expected {expected_str}, found [{result_clean}]")

        if correct_count == 8:
            print(f"PASS: Component 2 — All 8 values match expected calculations ({correct_count}/8) (0.4 pts)")
            total_score += 0.4
        elif correct_count > 0:
            partial = round(0.4 * (correct_count / 8), 2)
            print(f"PARTIAL: Component 2 — {correct_count}/8 values match ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No values match expected calculations (0/8)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Format consistency - all values have X.X% pattern (0.2 points)
    # Must have exactly one decimal place and a percent sign
    try:
        format_ok_count = 0
        pct_pattern = re.compile(r'^\d+\.\d%$')  # e.g., "93.8%" or "109.0%"
        for ri in range(1, 9):
            result_text = table.cell(ri, 4).text.strip()
            if pct_pattern.match(result_text):
                format_ok_count += 1
            else:
                print(f"  Row {ri}: Format check — [{result_text}] does not match X.X% pattern")

        if format_ok_count == 8:
            print(f"PASS: Component 3 — All 8 values have correct X.X% format ({format_ok_count}/8) (0.2 pts)")
            total_score += 0.2
        elif format_ok_count > 0:
            partial = round(0.2 * (format_ok_count / 8), 2)
            print(f"PARTIAL: Component 3 — {format_ok_count}/8 values have correct format ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No values have correct X.X% format (0/8)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice changes before verification
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
