"""
Initial Setup: Create a 10-page proposal document (Proposal_v2.docx) with realistic content
ending with a pricing table mid-page 10. No terms and conditions section.
Task ID: writer_pd_016
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_016'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_styled_paragraph(doc, text, font_name="Calibri", font_size=11, bold=False,
                          alignment=None, space_after=Pt(6), space_before=Pt(0)):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if alignment:
        para.paragraph_format.alignment = alignment
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = space_before
    return para


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ---- PAGE 1: Title Page ----
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run("Enterprise Digital Transformation")
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Strategic Implementation Proposal")
    run.font.size = Pt(18)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = info.add_run("Prepared for: Meridian Global Industries, Inc.\nPrepared by: Apex Consulting Group\nDate: March 15, 2025\nVersion: 2.0")
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ---- PAGE 2: Table of Contents ----
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary .................................................... 3",
        "2. Project Background ................................................... 4",
        "3. Scope of Work ........................................................... 5",
        "4. Methodology .............................................................. 6",
        "5. Project Timeline ........................................................ 7",
        "6. Team Structure .......................................................... 7",
        "7. Risk Assessment ........................................................ 8",
        "8. Deliverables ............................................................... 9",
        "9. Budget & Pricing ....................................................... 10",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        r = p.add_run(item)
        r.font.size = Pt(11)
        r.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ---- PAGE 3: Executive Summary ----
    doc.add_heading("1. Executive Summary", level=1)
    add_styled_paragraph(doc, (
        "Apex Consulting Group is pleased to present this comprehensive proposal for the "
        "digital transformation initiative at Meridian Global Industries. This engagement will "
        "modernize your enterprise infrastructure, streamline business processes, and establish "
        "a robust data analytics platform that positions Meridian for sustained competitive advantage "
        "in the evolving marketplace."
    ))
    add_styled_paragraph(doc, (
        "Our approach leverages proven methodologies refined across 15 years of enterprise consulting, "
        "combined with cutting-edge cloud-native architectures and AI-driven automation. We have "
        "successfully delivered similar transformations for Fortune 500 clients including Harlow "
        "Technologies, Brighton Manufacturing, and Cascade Financial Services."
    ))
    add_styled_paragraph(doc, (
        "The proposed 18-month engagement encompasses four primary workstreams: legacy system migration, "
        "process automation, data warehouse modernization, and change management. Our team of 24 specialists "
        "will work in close collaboration with your internal IT department and business stakeholders to ensure "
        "seamless integration and knowledge transfer throughout the project lifecycle."
    ))
    add_styled_paragraph(doc, (
        "Key expected outcomes include a 35% reduction in operational costs, 60% improvement in data processing "
        "speed, and a unified enterprise platform that eliminates the current fragmentation across 12 legacy systems. "
        "We project full return on investment within 30 months of project completion."
    ))

    doc.add_page_break()

    # ---- PAGE 4: Project Background ----
    doc.add_heading("2. Project Background", level=1)
    add_styled_paragraph(doc, (
        "Meridian Global Industries currently operates on a complex technology landscape comprising 12 "
        "interconnected legacy systems, several of which date back to the early 2000s. The existing ERP "
        "platform, originally deployed in 2004, has been incrementally modified over two decades, resulting "
        "in significant technical debt and reduced agility."
    ))
    add_styled_paragraph(doc, (
        "During our initial discovery phase conducted in January 2025, we identified several critical "
        "challenges affecting operational efficiency. The average data reconciliation cycle takes 72 hours "
        "across three departments, manual data entry accounts for approximately 2,400 staff-hours per month, "
        "and the current reporting infrastructure cannot deliver real-time analytics to decision-makers."
    ))
    add_styled_paragraph(doc, (
        "Additionally, Meridian's recent acquisition of Thornfield Distribution has introduced integration "
        "requirements that the current architecture cannot accommodate without substantial modification. "
        "The Thornfield systems operate on entirely different technology stacks, creating data silos that "
        "impede unified business intelligence across the combined entity."
    ))
    add_styled_paragraph(doc, (
        "Market analysis indicates that Meridian's key competitors, including Vanguard Solutions and Prism "
        "Industries, have already completed or initiated similar transformation programs. Delays in modernization "
        "risk eroding Meridian's market position, particularly in the Asia-Pacific region where digital-first "
        "competitors are gaining share rapidly."
    ))

    doc.add_page_break()

    # ---- PAGE 5: Scope of Work ----
    doc.add_heading("3. Scope of Work", level=1)
    add_styled_paragraph(doc, (
        "The scope of this engagement covers four integrated workstreams designed to deliver a comprehensive "
        "transformation of Meridian's digital infrastructure and operational processes."
    ))

    doc.add_heading("3.1 Legacy System Migration", level=2)
    add_styled_paragraph(doc, (
        "Migration of 12 legacy applications to a unified cloud-native platform built on Microsoft Azure. "
        "This includes the SAP ECC to SAP S/4HANA migration, retirement of the proprietary Meridian Inventory "
        "Management System (MIMS), and consolidation of three separate CRM instances into Salesforce Enterprise."
    ))

    doc.add_heading("3.2 Process Automation", level=2)
    add_styled_paragraph(doc, (
        "Implementation of robotic process automation (RPA) using UiPath for 45 identified manual processes "
        "across finance, procurement, and human resources. Priority automation targets include invoice processing, "
        "vendor onboarding, employee time tracking reconciliation, and quarterly compliance reporting."
    ))

    doc.add_heading("3.3 Data Warehouse Modernization", level=2)
    add_styled_paragraph(doc, (
        "Deployment of a modern data lakehouse architecture using Databricks on Azure, replacing the existing "
        "on-premises Oracle data warehouse. The new platform will support real-time streaming analytics, "
        "machine learning workloads, and self-service business intelligence through Power BI dashboards."
    ))

    doc.add_heading("3.4 Change Management", level=2)
    add_styled_paragraph(doc, (
        "A structured change management program encompassing stakeholder engagement, training curriculum "
        "development, communication planning, and adoption tracking. We will deploy our proprietary Apex "
        "Change Readiness Framework to ensure organizational alignment and minimize productivity disruptions."
    ))

    doc.add_page_break()

    # ---- PAGE 6: Methodology ----
    doc.add_heading("4. Methodology", level=1)
    add_styled_paragraph(doc, (
        "Apex Consulting employs a hybrid agile methodology that combines the governance rigor of traditional "
        "waterfall project management with the iterative flexibility of agile delivery. This approach, which we "
        "call Apex Adaptive Delivery (AAD), has been refined across 200+ enterprise engagements."
    ))
    add_styled_paragraph(doc, (
        "Phase 1 - Discovery & Planning (Months 1-2): Comprehensive assessment of current state, detailed "
        "requirements gathering through stakeholder workshops, architecture design, and creation of the project "
        "blueprint. Deliverables include the Solution Architecture Document, Migration Runbook, and Risk Registry."
    ))
    add_styled_paragraph(doc, (
        "Phase 2 - Foundation & Build (Months 3-8): Core platform provisioning, development of integration "
        "layers, RPA bot development, and data pipeline construction. Two-week sprint cycles with fortnightly "
        "stakeholder demonstrations ensure continuous alignment with business objectives."
    ))
    add_styled_paragraph(doc, (
        "Phase 3 - Migration & Testing (Months 9-14): Phased migration of legacy systems with parallel "
        "running periods, comprehensive UAT cycles, performance testing, and security validation. Each system "
        "migration follows our five-gate quality assurance process."
    ))
    add_styled_paragraph(doc, (
        "Phase 4 - Stabilization & Handover (Months 15-18): Hypercare support, knowledge transfer sessions, "
        "documentation finalization, and formal project closure. Includes 90-day post-go-live warranty period."
    ))

    doc.add_page_break()

    # ---- PAGE 7: Timeline & Team ----
    doc.add_heading("5. Project Timeline", level=1)
    add_styled_paragraph(doc, (
        "The following timeline outlines key milestones across the 18-month engagement:"
    ))

    timeline_table = doc.add_table(rows=7, cols=3)
    timeline_table.style = "Table Grid"
    headers = ["Milestone", "Target Date", "Status"]
    for i, h in enumerate(headers):
        cell = timeline_table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    timeline_data = [
        ["Project Kickoff", "April 1, 2025", "Pending"],
        ["Discovery Complete", "May 31, 2025", "Pending"],
        ["Platform Foundation Ready", "October 31, 2025", "Pending"],
        ["First System Migration", "January 15, 2026", "Pending"],
        ["Full Migration Complete", "June 30, 2026", "Pending"],
        ["Project Closure", "September 30, 2026", "Pending"],
    ]
    for r, row_data in enumerate(timeline_data, 1):
        for c, val in enumerate(row_data):
            cell = timeline_table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    doc.add_paragraph()

    doc.add_heading("6. Team Structure", level=1)
    add_styled_paragraph(doc, (
        "Our dedicated project team consists of 24 professionals with deep expertise in enterprise "
        "transformation, cloud architecture, and change management."
    ))

    team_table = doc.add_table(rows=7, cols=4)
    team_table.style = "Table Grid"
    team_headers = ["Name", "Role", "Experience", "Allocation"]
    for i, h in enumerate(team_headers):
        cell = team_table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    team_data = [
        ["Dr. Rachel Whitfield", "Engagement Partner", "22 years", "25%"],
        ["James Nakamura", "Program Director", "18 years", "100%"],
        ["Priya Kapoor", "Solution Architect", "15 years", "100%"],
        ["Chen Wei", "Data Engineering Lead", "12 years", "100%"],
        ["Sofia Andersson", "Change Management Lead", "14 years", "75%"],
        ["Marcus Thompson", "RPA Development Lead", "10 years", "100%"],
    ]
    for r, row_data in enumerate(team_data, 1):
        for c, val in enumerate(row_data):
            cell = team_table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    doc.add_page_break()

    # ---- PAGE 8: Risk Assessment ----
    doc.add_heading("7. Risk Assessment", level=1)
    add_styled_paragraph(doc, (
        "Our risk assessment methodology identifies, quantifies, and establishes mitigation strategies for "
        "all material project risks. The following table summarizes the top risks identified during our "
        "preliminary assessment."
    ))

    risk_table = doc.add_table(rows=7, cols=4)
    risk_table.style = "Table Grid"
    risk_headers = ["Risk", "Probability", "Impact", "Mitigation"]
    for i, h in enumerate(risk_headers):
        cell = risk_table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    risk_data = [
        ["Legacy data quality issues", "High", "High", "Pre-migration data cleansing sprints"],
        ["Stakeholder resistance to change", "Medium", "High", "Early engagement and champion network"],
        ["Integration complexity exceeds estimates", "Medium", "Medium", "Architecture spike sprints and buffer allocation"],
        ["Key personnel attrition", "Low", "High", "Cross-training and knowledge documentation"],
        ["Vendor platform updates during migration", "Medium", "Medium", "Version pinning and change freeze windows"],
        ["Regulatory compliance gaps", "Low", "High", "Dedicated compliance review at each gate"],
    ]
    for r, row_data in enumerate(risk_data, 1):
        for c, val in enumerate(row_data):
            cell = risk_table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    add_styled_paragraph(doc, (
        "Each risk is assigned a dedicated owner from the project team who is responsible for ongoing "
        "monitoring and activation of mitigation plans. Risk status is reported weekly in the project "
        "steering committee meeting and escalated to the executive sponsor when risk ratings change."
    ))
    add_styled_paragraph(doc, (
        "Additionally, we maintain a comprehensive risk register that is updated bi-weekly and shared with "
        "all stakeholders through the project portal. Historical risk data from comparable engagements "
        "informs our probability assessments and ensures realistic contingency planning."
    ))

    doc.add_page_break()

    # ---- PAGE 9: Deliverables ----
    doc.add_heading("8. Deliverables", level=1)
    add_styled_paragraph(doc, (
        "The following deliverables will be produced throughout the engagement lifecycle. Each deliverable "
        "undergoes formal quality review and requires sign-off from the designated Meridian stakeholder."
    ))

    deliverables_table = doc.add_table(rows=11, cols=3)
    deliverables_table.style = "Table Grid"
    del_headers = ["Deliverable", "Phase", "Acceptance Criteria"]
    for i, h in enumerate(del_headers):
        cell = deliverables_table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    del_data = [
        ["Solution Architecture Document", "Phase 1", "Approved by CTO and Solution Architect"],
        ["Migration Runbook", "Phase 1", "Approved by IT Operations Manager"],
        ["Risk Registry", "Phase 1", "Approved by Program Director"],
        ["Cloud Platform Foundation", "Phase 2", "Passes security audit and performance benchmarks"],
        ["RPA Bot Suite (45 bots)", "Phase 2", "95% accuracy in UAT across all processes"],
        ["Data Lakehouse Platform", "Phase 2", "Passes data integrity validation"],
        ["Migrated SAP S/4HANA System", "Phase 3", "Passes parallel run with zero critical defects"],
        ["Unified CRM Platform", "Phase 3", "User acceptance sign-off from sales leadership"],
        ["Training Materials Package", "Phase 4", "Approved by Change Management Lead"],
        ["Operations Handbook", "Phase 4", "Approved by IT Operations and Business Owners"],
    ]
    for r, row_data in enumerate(del_data, 1):
        for c, val in enumerate(row_data):
            cell = deliverables_table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    add_styled_paragraph(doc, (
        "Formal deliverable reviews follow the Apex Quality Gate process, which requires sign-off from "
        "both the project team and the client stakeholder before a deliverable is marked as accepted. "
        "Any rejected deliverables enter a remediation cycle with clearly defined resolution timelines."
    ))

    doc.add_page_break()

    # ---- PAGE 10: Budget & Pricing (ends mid-page) ----
    doc.add_heading("9. Budget & Pricing", level=1)
    add_styled_paragraph(doc, (
        "The following pricing structure reflects our commitment to transparent, value-based engagement "
        "economics. All fees are quoted in USD and are exclusive of applicable taxes."
    ))

    pricing_table = doc.add_table(rows=8, cols=3)
    pricing_table.style = "Table Grid"
    price_headers = ["Work Package", "Duration", "Investment (USD)"]
    for i, h in enumerate(price_headers):
        cell = pricing_table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    price_data = [
        ["Phase 1: Discovery & Planning", "2 months", "$285,000"],
        ["Phase 2: Foundation & Build", "6 months", "$1,420,000"],
        ["Phase 3: Migration & Testing", "6 months", "$1,680,000"],
        ["Phase 4: Stabilization & Handover", "4 months", "$520,000"],
        ["Change Management Program", "18 months", "$345,000"],
        ["Project Management Office", "18 months", "$480,000"],
        ["Total Investment", "18 months", "$4,730,000"],
    ]
    for r, row_data in enumerate(price_data, 1):
        for c, val in enumerate(row_data):
            cell = pricing_table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if r == 7:  # Total row
                run.bold = True

    add_styled_paragraph(doc, (
        "Payment terms: 30% upon contract execution, followed by monthly invoicing based on milestone "
        "completion as outlined in the project plan. All invoices are net 30 days."
    ))
    add_styled_paragraph(doc, (
        "This investment includes all professional services, travel expenses within the continental United States, "
        "project management tools and licenses, and the 90-day post-go-live warranty period. Third-party software "
        "licensing costs (Azure, Salesforce, UiPath, Databricks) are excluded and will be procured directly by "
        "Meridian under our negotiated enterprise agreements."
    ))

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
