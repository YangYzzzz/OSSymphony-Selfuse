"""
Reward Script: Add Terms and Conditions section to Proposal_v2.docx
Task ID: writer_pd_016
Domain: libreoffice_writer
Scoring:
  Component 1: "Terms and Conditions" heading in Heading 1 style (0.20)
  Component 2: Page break before T&C section (0.15)
  Component 3: Horizontal rule (paragraph border) before the heading (0.15)
  Component 4: 10 numbered clauses present (0.20)
  Component 5: Clause numbers are bold (0.10)
  Component 6: Clause titles are bold italic (0.10)
  Component 7: Body text is regular 10pt (0.10)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_016'


def persist_app_state(domain):
    """Best-effort save if LibreOffice is open."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the "Terms and Conditions" heading paragraph
    tc_heading_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if text == 'terms and conditions':
            tc_heading_idx = i
            break

    # Component 1: "Terms and Conditions" heading in Heading 1 style (0.20 points)
    try:
        if tc_heading_idx is not None:
            heading_para = doc.paragraphs[tc_heading_idx]
            if heading_para.style.name == 'Heading 1':
                print(f"PASS: Component 1 - 'Terms and Conditions' heading found at para {tc_heading_idx} in Heading 1 style (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 1 - 'Terms and Conditions' found but style is '{heading_para.style.name}', expected 'Heading 1'")
        else:
            print("FAIL: Component 1 - No 'Terms and Conditions' heading found in document")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # If no T&C heading found, remaining checks will fail — exit early
    if tc_heading_idx is None:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Component 2: Page break before T&C section (0.15 points)
    # The page break should be in one of the paragraphs before the T&C heading
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        found_page_break = False
        # Check paragraphs from tc_heading_idx-3 to tc_heading_idx (the break could be
        # in the heading itself or in an empty paragraph before it)
        search_start = max(0, tc_heading_idx - 5)
        for idx in range(search_start, tc_heading_idx + 1):
            para = doc.paragraphs[idx]
            # Check page_break_before property
            if para.paragraph_format.page_break_before:
                found_page_break = True
                break
            # Check for w:br type=page in runs
            for run in para.runs:
                for br in run.element.findall('.//w:br', ns):
                    btype = br.attrib.get(qn('w:type'))
                    if btype == 'page':
                        found_page_break = True
                        break
                if found_page_break:
                    break
            if found_page_break:
                break

        if found_page_break:
            print(f"PASS: Component 2 - Page break found before T&C section (0.15 pts)")
            total_score += 0.15
        else:
            print("FAIL: Component 2 - No page break found before T&C section")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Horizontal rule (paragraph border) before the heading (0.15 points)
    try:
        found_hr = False
        # Check paragraphs between page break area and the heading
        for idx in range(search_start, tc_heading_idx):
            para = doc.paragraphs[idx]
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    # Check for bottom border (common HR representation)
                    bottom = pBdr.find(qn('w:bottom'))
                    top = pBdr.find(qn('w:top'))
                    if bottom is not None or top is not None:
                        found_hr = True
                        border_el = bottom if bottom is not None else top
                        print(f"PASS: Component 3 - Horizontal rule (paragraph border) found at para {idx} (0.15 pts)")
                        total_score += 0.15
                        break

        if not found_hr:
            print("FAIL: Component 3 - No horizontal rule (paragraph border) found before T&C heading")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Collect clause paragraphs (paragraphs after T&C heading that start with a number)
    clause_paras = []
    for idx in range(tc_heading_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[idx]
        text = para.text.strip()
        if not text:
            continue
        # Match numbered clauses like "1. ...", "2. ...", etc.
        match = re.match(r'^(\d+)\.\s+', text)
        if match:
            clause_paras.append((int(match.group(1)), idx, para))

    # Component 4: 10 numbered clauses present (0.20 points)
    try:
        clause_numbers = [c[0] for c in clause_paras]
        if len(clause_paras) >= 10 and all(n in clause_numbers for n in range(1, 11)):
            print(f"PASS: Component 4 - Found {len(clause_paras)} numbered clauses (1-10) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 - Expected 10 clauses numbered 1-10, found {len(clause_paras)} clauses with numbers {clause_numbers}")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Components 5, 6, 7: Check formatting of clauses
    # Each clause should have: bold number, bold italic title, regular 10pt body
    # Structure per clause paragraph: runs like [bold "1. "][bold+italic "Title: "][regular 10pt "body text..."]

    bold_number_count = 0
    bold_italic_title_count = 0
    regular_body_count = 0
    clauses_checked = min(len(clause_paras), 10)

    for clause_num, idx, para in clause_paras[:10]:
        try:
            runs = para.runs
            if len(runs) < 2:
                print(f"  Clause {clause_num}: Only {len(runs)} run(s), expected >= 3")
                continue

            # Run 0: clause number (e.g., "1. ") - should be bold
            number_run = runs[0]
            if number_run.font.bold is True or number_run.bold is True:
                bold_number_count += 1
            else:
                print(f"  Clause {clause_num}: Number run bold={number_run.font.bold}, expected True")

            # Run 1: clause title (e.g., "Acceptance of Terms: ") - should be bold+italic
            if len(runs) >= 2:
                title_run = runs[1]
                is_bold = title_run.font.bold is True or title_run.bold is True
                is_italic = title_run.font.italic is True or title_run.italic is True
                if is_bold and is_italic:
                    bold_italic_title_count += 1
                else:
                    print(f"  Clause {clause_num}: Title run bold={title_run.font.bold}, italic={title_run.font.italic}")

            # Run 2+: body text - should NOT be bold, NOT italic, and 10pt
            if len(runs) >= 3:
                body_run = runs[2]
                is_not_bold = body_run.font.bold is False or body_run.font.bold is None
                is_not_italic = body_run.font.italic is False or body_run.font.italic is None
                # Check font size: 10pt = 127000 EMU
                size_ok = False
                if body_run.font.size is not None:
                    size_pt = body_run.font.size.pt
                    size_ok = abs(size_pt - 10.0) < 0.5
                if is_not_bold and is_not_italic and size_ok:
                    regular_body_count += 1
                else:
                    size_info = body_run.font.size.pt if body_run.font.size else 'None'
                    print(f"  Clause {clause_num}: Body run bold={body_run.font.bold}, italic={body_run.font.italic}, size={size_info}pt")
        except Exception as e:
            print(f"  Clause {clause_num}: ERROR checking formatting - {e}")

    # Component 5: Clause numbers are bold (0.10 points)
    try:
        if clauses_checked > 0 and bold_number_count >= clauses_checked * 0.8:
            print(f"PASS: Component 5 - {bold_number_count}/{clauses_checked} clause numbers are bold (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 - Only {bold_number_count}/{clauses_checked} clause numbers are bold")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: Clause titles are bold italic (0.10 points)
    try:
        if clauses_checked > 0 and bold_italic_title_count >= clauses_checked * 0.8:
            print(f"PASS: Component 6 - {bold_italic_title_count}/{clauses_checked} clause titles are bold italic (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 - Only {bold_italic_title_count}/{clauses_checked} clause titles are bold italic")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    # Component 7: Body text is regular 10pt (0.10 points)
    try:
        if clauses_checked > 0 and regular_body_count >= clauses_checked * 0.8:
            print(f"PASS: Component 7 - {regular_body_count}/{clauses_checked} clause body texts are regular 10pt (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 7 - Only {regular_body_count}/{clauses_checked} clause body texts are regular 10pt")
    except Exception as e:
        print(f"ERROR: Component 7 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
