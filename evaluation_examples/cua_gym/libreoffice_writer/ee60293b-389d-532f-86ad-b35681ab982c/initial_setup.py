"""
Initial Setup: Corporate governance policy document with unnumbered headings
Task ID: writer_lec_029
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_029'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading('Corporate Governance Policy Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        'This document establishes the corporate governance framework for Meridian '
        'Holdings International. It outlines the principles, structures, and processes '
        'by which the company is directed, managed, and held accountable to stakeholders. '
        'All employees, officers, and directors are expected to comply with the policies '
        'set forth herein.'
    )

    # === SECTION 1: Board of Directors ===
    doc.add_heading('Board of Directors', level=1)
    doc.add_paragraph(
        'The Board of Directors serves as the primary governing body of Meridian Holdings '
        'International. The Board is responsible for providing strategic direction, overseeing '
        'management performance, and ensuring that the company operates in the best interests '
        'of all stakeholders.'
    )

    # -- 1.1 Composition and Structure --
    doc.add_heading('Composition and Structure', level=2)
    doc.add_paragraph(
        'The Board shall consist of no fewer than seven and no more than fifteen members, '
        'with at least two-thirds qualifying as independent directors under applicable '
        'securities regulations. Board members serve staggered three-year terms to ensure '
        'continuity of governance.'
    )

    # -- 1.1.1 Qualifications --
    doc.add_heading('Director Qualifications', level=3)
    doc.add_paragraph(
        'Each director candidate must demonstrate expertise in at least one of the following '
        'areas: corporate finance, risk management, industry-specific knowledge, legal and '
        'regulatory compliance, or information technology. The Nominating Committee evaluates '
        'candidates based on a competency matrix reviewed annually.'
    )

    # -- 1.1.1.a) --
    doc.add_heading('Independence Standards', level=4)
    doc.add_paragraph(
        'A director is considered independent if they have no material relationship with the '
        'company, either directly or as a partner, shareholder, or officer of an organization '
        'that has a relationship with the company. Material relationships include any financial '
        'transaction exceeding $120,000 annually within the preceding three fiscal years.'
    )

    doc.add_heading('Diversity Requirements', level=4)
    doc.add_paragraph(
        'The Board is committed to maintaining diversity across gender, ethnicity, professional '
        'background, and geographic representation. The target is to achieve at least 40% '
        'representation of underrepresented groups by December 2026.'
    )

    # -- 1.1.2 --
    doc.add_heading('Board Succession Planning', level=3)
    doc.add_paragraph(
        'The Nominating and Governance Committee maintains a rolling three-year succession '
        'plan for all Board positions. Emergency succession protocols are reviewed semi-annually '
        'and include provisions for interim appointments when vacancies arise unexpectedly.'
    )

    doc.add_heading('Emergency Vacancy Procedures', level=4)
    doc.add_paragraph(
        'In the event of an unplanned vacancy, the Lead Independent Director shall convene '
        'a special session within five business days. The remaining directors may appoint an '
        'interim member by a two-thirds supermajority vote, subject to ratification at the '
        'next annual general meeting.'
    )

    # -- 1.2 Meetings and Procedures --
    doc.add_heading('Meetings and Procedures', level=2)
    doc.add_paragraph(
        'The Board shall hold a minimum of six regularly scheduled meetings per fiscal year, '
        'with additional special meetings convened as necessary. A quorum of at least 60% of '
        'serving directors is required for the transaction of business.'
    )

    doc.add_heading('Agenda and Materials', level=3)
    doc.add_paragraph(
        'The Corporate Secretary distributes meeting agendas and supporting materials no fewer '
        'than seven calendar days prior to each scheduled Board meeting. Directors may request '
        'additional agenda items through the Chairperson or Lead Independent Director at least '
        'ten days before the meeting date.'
    )

    doc.add_heading('Standard Agenda Items', level=4)
    doc.add_paragraph(
        'Each regular meeting shall include: approval of prior meeting minutes, CEO operational '
        'report, CFO financial update, committee reports, old business, new business, and an '
        'executive session without management present. Special agenda items related to mergers, '
        'acquisitions, or capital expenditures exceeding $5 million require separate approval.'
    )

    doc.add_heading('Voting Procedures', level=3)
    doc.add_paragraph(
        'All Board resolutions require a simple majority vote unless otherwise specified in '
        'the Articles of Incorporation. Proxy voting is permitted only with written authorization '
        'submitted to the Corporate Secretary at least 48 hours before the meeting.'
    )

    # === SECTION 2: Executive Management ===
    doc.add_heading('Executive Management', level=1)
    doc.add_paragraph(
        'The executive management team is responsible for the day-to-day operations of the '
        'company and for implementing the strategic direction established by the Board of '
        'Directors. All executive officers report to the Chief Executive Officer, who in turn '
        'reports to the Board.'
    )

    doc.add_heading('Chief Executive Officer', level=2)
    doc.add_paragraph(
        'The CEO has authority over all operational decisions within the parameters set by '
        'the Board. The CEO may delegate authority to other executives but retains ultimate '
        'responsibility for company performance. Annual performance evaluations are conducted '
        'by the Compensation Committee using pre-established quantitative and qualitative metrics.'
    )

    doc.add_heading('Authority Limits', level=3)
    doc.add_paragraph(
        'The CEO may approve capital expenditures up to $2 million per transaction and '
        'operating expenses within the approved annual budget. Expenditures exceeding these '
        'thresholds require Board approval. The CEO may not enter into contracts with a term '
        'exceeding five years without prior Board authorization.'
    )

    doc.add_heading('Signing Authority', level=4)
    doc.add_paragraph(
        'The CEO holds first-level signing authority for all corporate obligations. For '
        'obligations exceeding $500,000, dual signatures are required from the CEO and either '
        'the CFO or the General Counsel. All real estate transactions require Board resolution '
        'regardless of value.'
    )

    doc.add_heading('Chief Financial Officer', level=2)
    doc.add_paragraph(
        'The CFO oversees all financial operations, including treasury management, financial '
        'reporting, tax compliance, and investor relations. The CFO provides monthly financial '
        'reports to the Board and certifies the accuracy of all public financial disclosures '
        'in accordance with applicable securities laws.'
    )

    doc.add_heading('Financial Reporting Standards', level=3)
    doc.add_paragraph(
        'All financial statements are prepared in accordance with International Financial '
        'Reporting Standards (IFRS) as adopted by the jurisdictions in which the company '
        'operates. Management Discussion and Analysis sections must be reviewed by the Audit '
        'Committee prior to publication.'
    )

    doc.add_heading('Internal Controls Certification', level=4)
    doc.add_paragraph(
        'The CFO certifies quarterly to the Audit Committee that all internal financial '
        'controls are operating effectively. Any material weaknesses identified must be '
        'reported to the Board within five business days of discovery and remediated within '
        '90 calendar days.'
    )

    # === SECTION 3: Risk Management ===
    doc.add_heading('Risk Management Framework', level=1)
    doc.add_paragraph(
        'Meridian Holdings maintains a comprehensive enterprise risk management framework '
        'aligned with the COSO Enterprise Risk Management framework. The Chief Risk Officer '
        'reports directly to the Board Risk Committee and has unfettered access to all company '
        'operations and records.'
    )

    doc.add_heading('Risk Identification and Assessment', level=2)
    doc.add_paragraph(
        'Risk identification is performed on a continuous basis using a combination of '
        'top-down strategic risk assessment and bottom-up operational risk reporting. '
        'All identified risks are categorized by likelihood and impact using a standardized '
        '5x5 risk matrix and are recorded in the Enterprise Risk Register.'
    )

    doc.add_heading('Risk Categories', level=3)
    doc.add_paragraph(
        'The Enterprise Risk Register maintains risks across five primary categories: '
        'strategic, operational, financial, compliance, and reputational. Each category '
        'has designated risk owners at the senior vice president level or above.'
    )

    doc.add_heading('Emerging Risk Monitoring', level=4)
    doc.add_paragraph(
        'The Risk Intelligence Unit monitors geopolitical developments, regulatory changes, '
        'technology disruptions, and market shifts that could create new risk exposures. '
        'Emerging risk reports are presented to the Board Risk Committee on a quarterly basis, '
        'with ad hoc alerts issued for significant developments.'
    )

    doc.add_heading('Risk Mitigation Strategies', level=2)
    doc.add_paragraph(
        'For each identified risk rated medium or higher on the risk matrix, a formal '
        'mitigation plan must be developed and approved by the responsible executive. '
        'Mitigation strategies include risk avoidance, risk transfer through insurance, '
        'risk reduction through controls, and risk acceptance with enhanced monitoring.'
    )

    doc.add_heading('Insurance Coverage Requirements', level=3)
    doc.add_paragraph(
        'The company maintains comprehensive insurance coverage including Directors and '
        'Officers liability, professional indemnity, cyber liability, and general commercial '
        'liability policies. Coverage limits are reviewed annually and must meet or exceed '
        'industry benchmarks for companies of comparable size and risk profile.'
    )

    doc.add_heading('Claims Reporting Procedures', level=4)
    doc.add_paragraph(
        'All potential claims must be reported to the Legal Department within 24 hours of '
        'discovery. The General Counsel notifies relevant insurers within the timeframes '
        'specified in each policy. Failure to report claims promptly may result in coverage '
        'denial and personal liability for the responsible officer.'
    )

    # === SECTION 4: Compliance and Ethics ===
    doc.add_heading('Compliance and Ethics Program', level=1)
    doc.add_paragraph(
        'The company maintains a comprehensive compliance and ethics program designed to '
        'prevent, detect, and respond to violations of law, regulation, and internal policy. '
        'The Chief Compliance Officer reports to the Audit Committee and has direct access '
        'to the Board without management intermediation.'
    )

    doc.add_heading('Code of Conduct', level=2)
    doc.add_paragraph(
        'Every employee, contractor, officer, and director must acknowledge the Code of '
        'Conduct annually. The Code addresses conflicts of interest, gifts and entertainment, '
        'anti-bribery compliance, data privacy, insider trading prohibitions, and standards '
        'for workplace behavior.'
    )

    doc.add_heading('Whistleblower Protection', level=3)
    doc.add_paragraph(
        'The company operates an anonymous reporting hotline managed by an independent '
        'third-party provider. All reports are investigated by the Ethics and Compliance '
        'team, with oversight from the Audit Committee for matters involving senior '
        'management. Retaliation against good-faith reporters is strictly prohibited and '
        'constitutes grounds for immediate termination.'
    )

    doc.add_heading('Investigation Procedures', level=4)
    doc.add_paragraph(
        'Upon receipt of a whistleblower report, the Ethics and Compliance team initiates '
        'a preliminary assessment within 48 hours. Full investigations are completed within '
        '60 calendar days unless complexity requires an extension approved by the Chief '
        'Compliance Officer. Investigation findings are reported to the Audit Committee '
        'at its next scheduled meeting.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
