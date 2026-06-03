"""
Initial Setup: AI Literature Review with bibliography placeholders
Task ID: osworld_writer_biblio_010
Domain: libreoffice_writer

Creates ai_literature_review.docx with:
- 5 sections plus References section
- Conclusion paragraph with <cite A> and <cite B> placeholders
- 10 existing numbered APA references
- Opens the file in LibreOffice Writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_biblio_010'
OUTPUT = f'{WORKDIR}/ai_literature_review.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Advances in Large-Scale Language Models: A Literature Review", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Author info
    author_para = doc.add_paragraph("Department of Computer Science, Stanford University")
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.add_run("\nsubmitted@cs.stanford.edu").italic = True

    doc.add_paragraph()  # spacing

    # --- Abstract ---
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This literature review examines the rapid evolution of large-scale language models over the past decade. "
        "We survey key architectures, training methodologies, and emergent capabilities that have defined the field. "
        "Special attention is given to transformer-based models, self-supervised pre-training, and the scaling laws "
        "that govern model performance. Our analysis synthesizes findings from over 50 research publications to "
        "identify major trends, open challenges, and promising directions for future inquiry."
    )

    # --- 1. Introduction ---
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "The field of natural language processing (NLP) has undergone a paradigm shift with the advent of "
        "large-scale pre-trained language models. Beginning with word embeddings (Mikolov et al., 2013) and "
        "progressing through ELMo (Peters et al., 2018) and BERT (Devlin et al., 2019), the community has "
        "witnessed increasingly capable models that can be fine-tuned for a wide variety of downstream tasks "
        "(1, 2, 3)."
    )
    doc.add_paragraph(
        "The introduction of the transformer architecture by Vaswani et al. (2017) marked a pivotal turning "
        "point, enabling parallelized training and attention-based context modeling at unprecedented scales. "
        "Subsequent work demonstrated that models trained on sufficiently large corpora could generalize "
        "remarkably well across domains without task-specific fine-tuning (4, 5)."
    )
    doc.add_paragraph(
        "This review is organized as follows: Section 2 surveys foundational work in neural language modeling; "
        "Section 3 examines the transformer era; Section 4 discusses emergent capabilities and scaling; "
        "Section 5 addresses societal implications; and Section 6 concludes with open challenges."
    )

    # --- 2. Foundational Language Models ---
    doc.add_heading("2. Foundational Language Models", level=1)
    doc.add_paragraph(
        "Early neural language models used recurrent architectures to capture sequential dependencies in text. "
        "Long short-term memory networks (LSTMs) (Hochreiter & Schmidhuber, 1997) and gated recurrent units "
        "(GRUs) (Cho et al., 2014) achieved strong performance on tasks such as machine translation and "
        "sentiment analysis (6). However, these architectures struggled with long-range dependencies and "
        "were computationally expensive to train."
    )
    doc.add_paragraph(
        "The introduction of contextualized word representations through ELMo (Peters et al., 2018) "
        "demonstrated that word meaning could be disambiguated using bidirectional language model pre-training. "
        "This approach yielded substantial improvements across six NLP benchmarks, establishing the paradigm "
        "of pre-train-then-fine-tune that dominates the field today (7)."
    )

    # --- 3. The Transformer Era ---
    doc.add_heading("3. The Transformer Era and BERT-Family Models", level=1)
    doc.add_paragraph(
        "The transformer architecture introduced by Vaswani et al. (2017) relies entirely on self-attention "
        "mechanisms, dispensing with recurrence and convolution. This design choice enables massive "
        "parallelization during training and allows the model to directly attend to any position in the "
        "input sequence, regardless of distance (4)."
    )
    doc.add_paragraph(
        "BERT (Bidirectional Encoder Representations from Transformers) by Devlin et al. (2019) applied "
        "the transformer encoder to bidirectional pre-training using masked language modeling and "
        "next-sentence prediction objectives. The model achieved state-of-the-art results on eleven NLP "
        "benchmarks upon its release, demonstrating the power of large-scale pre-training (2)."
    )
    doc.add_paragraph(
        "Subsequent BERT variants addressed various limitations: RoBERTa (Liu et al., 2019) improved "
        "training with dynamic masking and larger batches; ALBERT (Lan et al., 2020) introduced "
        "parameter sharing for efficiency; and DistilBERT (Sanh et al., 2019) demonstrated knowledge "
        "distillation to compress models without significant performance loss (8, 9, 10)."
    )

    # --- 4. Scaling and Emergent Capabilities ---
    doc.add_heading("4. Scaling Laws and Emergent Capabilities", level=1)
    doc.add_paragraph(
        "A key empirical finding in recent years concerns the predictable relationship between model size, "
        "training compute, data volume, and downstream performance. Kaplan et al. (2020) established power-law "
        "scaling relationships, showing that model performance improves smoothly as scale increases along "
        "each of these dimensions (5)."
    )
    doc.add_paragraph(
        "More striking are reports of emergent capabilities—abilities that appear abruptly at scale thresholds "
        "and are absent in smaller models. These include multi-step arithmetic reasoning, chain-of-thought "
        "problem solving, and in-context learning from a handful of examples provided in the prompt. "
        "Such capabilities challenge traditional learning theory and raise new questions about the nature "
        "of intelligence in statistical models."
    )

    # --- 5. Societal Implications ---
    doc.add_heading("5. Societal Implications and Responsible Deployment", level=1)
    doc.add_paragraph(
        "The rapid proliferation of large language models has prompted significant discussion about their "
        "societal impacts. Issues include the potential for generating misinformation, perpetuating biases "
        "present in training data, and enabling malicious uses such as automated phishing or propaganda. "
        "Researchers and policymakers are actively developing frameworks for responsible deployment (3)."
    )
    doc.add_paragraph(
        "Environmental concerns have also garnered attention. Training very large models requires substantial "
        "computational resources, translating to significant energy consumption and carbon emissions. "
        "Efficiency research, model compression, and carbon-aware compute scheduling are active areas aimed "
        "at mitigating these impacts while preserving model capabilities."
    )

    # --- 6. Conclusion ---
    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(
        "This literature review has traced the development of large-scale language models from early "
        "recurrent architectures through the transformer revolution and the era of massive pre-training. "
        "The field has advanced remarkably quickly, with models demonstrating increasingly general capabilities "
        "across a broad spectrum of language understanding and generation tasks."
    )
    conclusion_para = doc.add_paragraph(
        "Looking forward, the community is exploring models of ever-greater scale and capability. "
        "The introduction of few-shot learning paradigms <cite A> and the analysis of foundation model "
        "opportunities and risks <cite B> represent two particularly influential recent contributions "
        "that will shape the trajectory of the field. Addressing open challenges in robustness, "
        "interpretability, and equitable access will be essential to realizing the full potential "
        "of this technology."
    )
    doc.add_paragraph(
        "We hope this review serves as a useful entry point for researchers entering the field and "
        "a consolidated reference for those seeking to situate their work within the broader landscape "
        "of large-scale language model research."
    )

    # --- References ---
    doc.add_heading("References", level=1)

    references = [
        "1. Mikolov, T., Chen, K., Corrado, G., & Dean, J. (2013). Efficient estimation of word representations in vector space. arXiv preprint arXiv:1301.3781.",
        "2. Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. Proceedings of NAACL-HLT 2019, 4171-4186.",
        "3. Bender, E. M., Gebru, T., McMillan-Major, A., & Shmitchell, S. (2021). On the dangers of stochastic parrots: Can language models be too big? Proceedings of FAccT 2021, 610-623.",
        "4. Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, L., & Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30.",
        "5. Kaplan, J., McCandlish, S., Henighan, T., Brown, T. B., Chess, B., Child, R., Gray, S., Radford, A., Wu, J., & Amodei, D. (2020). Scaling laws for neural language models. arXiv preprint arXiv:2001.08361.",
        "6. Cho, K., van Merri\u00ebnboer, B., Gulcehre, C., Bahdanau, D., Bougares, F., Schwenk, H., & Bengio, Y. (2014). Learning phrase representations using RNN encoder-decoder for statistical machine translation. Proceedings of EMNLP 2014, 1724-1734.",
        "7. Peters, M. E., Neumann, M., Iyyer, M., Gardner, M., Clark, C., Lee, K., & Zettlemoyer, L. (2018). Deep contextualized word representations. Proceedings of NAACL-HLT 2018, 2227-2237.",
        "8. Liu, Y., Ott, M., Goyal, N., Du, J., Joshi, M., Chen, D., Levy, O., Lewis, M., Zettlemoyer, L., & Stoyanov, V. (2019). RoBERTa: A robustly optimized BERT pretraining approach. arXiv preprint arXiv:1907.11692.",
        "9. Lan, Z., Chen, M., Goodman, S., Gimpel, K., Sharma, P., & Soricut, R. (2020). ALBERT: A lite BERT for self-supervised learning of language representations. Proceedings of ICLR 2020.",
        "10. Sanh, V., Debut, L., Chaumond, J., & Wolf, T. (2019). DistilBERT, a distilled version of BERT: smaller, faster, cheaper and lighter. arXiv preprint arXiv:1910.01108.",
    ]

    for ref in references:
        doc.add_paragraph(ref)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
