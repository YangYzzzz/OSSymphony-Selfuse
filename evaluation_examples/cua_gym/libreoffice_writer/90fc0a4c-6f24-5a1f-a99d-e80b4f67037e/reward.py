"""
Reward Script: Create a formatted financial report by transferring multiple ranges from
'financials.ods' into 'Q4_report.odt' with 3 tables, captions, and analysis paragraphs.
Task ID: osworld_multi_apps_doc_calc_to_writer_007
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): 3 tables present in Q4_report.odt
  Component 2 (0.25): 3 caption paragraphs present (Table 1/2/3 prefixed, bold)
  Component 3 (0.25): 3 analysis paragraphs (non-caption, non-empty text paragraphs)
  Component 4 (0.25): Table structure validity (correct row/col counts with bold headers)
"""

import os

# python-docx for reading .odt (which is actually OOXML format here)
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_doc_calc_to_writer_007'

DOC_PATH = f'{WORKDIR}/Documents/Q4_report.odt'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must exist and be loadable
    if not os.path.exists(file_path):
        print(f"CRITICAL: File not found: {file_path}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load document {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect basic document stats
    paragraphs = doc.paragraphs
    tables = doc.tables
    num_tables = len(tables)
    num_paras = len(paragraphs)

    print(f"Document stats: {num_paras} paragraphs, {num_tables} tables")

    # Component 1: 3 tables present in the document (0.25 points)
    # Initial doc has 0 tables; golden doc has 3 tables
    try:
        if num_tables >= 3:
            print(f"PASS: Component 1 — 3 tables present (found {num_tables}) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — expected 3 tables, found {num_tables}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 3 caption paragraphs present with 'Table 1', 'Table 2', 'Table 3' prefixes (0.25 points)
    # Each caption must start with "Table N:" and be in bold
    # Initial doc has 0 paragraphs; golden has captions
    try:
        captions_found = []
        for para in paragraphs:
            text = para.text.strip()
            for n in [1, 2, 3]:
                if text.startswith(f"Table {n}:"):
                    # Check if bold
                    is_bold = any(run.bold for run in para.runs if run.text.strip())
                    captions_found.append((n, text[:60], is_bold))
                    break

        caption_numbers = [c[0] for c in captions_found]
        bold_captions = [c for c in captions_found if c[2]]

        if len(captions_found) >= 3 and 1 in caption_numbers and 2 in caption_numbers and 3 in caption_numbers:
            if len(bold_captions) >= 3:
                print(f"PASS: Component 2 — all 3 captions present and bold (0.25 pts)")
                total_score += 0.25
            elif len(bold_captions) >= 2:
                print(f"PASS (partial): Component 2 — 3 captions present, {len(bold_captions)} are bold, awarding 0.15 pts")
                total_score += 0.15
            else:
                print(f"FAIL: Component 2 — 3 captions present but only {len(bold_captions)} are bold")
                print(f"  Captions: {captions_found}")
        elif len(captions_found) == 2:
            print(f"FAIL: Component 2 — only 2 of 3 captions found: {captions_found}")
        elif len(captions_found) == 1:
            print(f"FAIL: Component 2 — only 1 of 3 captions found: {captions_found}")
        else:
            print(f"FAIL: Component 2 — no captions found (Table 1/2/3 prefixes missing)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Analysis paragraphs present (non-caption, non-empty) (0.25 points)
    # The task requires a brief paragraph of analysis after each table.
    # We expect at least 3 non-caption paragraphs with meaningful text (>20 chars)
    # Initial doc has 0 paragraphs
    try:
        caption_prefixes = ("Table 1:", "Table 2:", "Table 3:")
        analysis_paras = []
        for para in paragraphs:
            text = para.text.strip()
            # Skip captions and empty paragraphs
            if not text:
                continue
            if any(text.startswith(p) for p in caption_prefixes):
                continue
            # Must be a substantive paragraph (>20 chars to avoid trivial entries)
            if len(text) > 20:
                analysis_paras.append(text[:80])

        if len(analysis_paras) >= 3:
            total_score += 0.25  # 3+ analysis paragraphs found
            print(f"PASS: Component 3 — {len(analysis_paras)} analysis paragraphs found (0.25 pts)")
            for ap in analysis_paras[:3]:
                print(f"  Analysis: {ap!r}")
        elif len(analysis_paras) == 2:
            print(f"PASS (partial): Component 3 — 2 of 3 analysis paragraphs found (0.15 pts)")
            total_score += 0.15
        elif len(analysis_paras) == 1:
            print(f"FAIL: Component 3 — only 1 analysis paragraph found: {analysis_paras}")
        else:
            print(f"FAIL: Component 3 — no analysis paragraphs found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Table structure validity — correct sizes and bold headers (0.25 points)
    # Expected:
    #   Table 0 (Revenue): 6 rows x 4 cols, bold headers
    #   Table 1 (Expenses): 5 rows x 4 cols, bold headers
    #   Table 2 (Profit Summary): 5 rows x 2 cols, bold headers
    # Initial doc has 0 tables, so this will fail on initial_env
    try:
        if num_tables >= 3:
            structure_checks = []

            # Expected table dimensions
            expected = [
                (6, 4, "Revenue"),     # Table 1
                (5, 4, "Expenses"),    # Table 2
                (5, 2, "Profit"),      # Table 3
            ]

            structure_fail_count = 0
            for t_idx in range(3):
                table = tables[t_idx]
                n_rows = len(table.rows)
                n_cols = len(table.columns)
                exp_rows, exp_cols, label = expected[t_idx]

                # Check dimensions
                dim_ok = (n_rows == exp_rows and n_cols == exp_cols)

                # Check header row bold
                header_row = table.rows[0]
                header_bold_count = 0
                for cell in header_row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip() and run.bold:
                                header_bold_count += 1
                                break
                header_bold_ok = header_bold_count >= 1

                structure_checks.append({
                    'label': label,
                    'dim_ok': dim_ok,
                    'actual_dims': (n_rows, n_cols),
                    'expected_dims': (exp_rows, exp_cols),
                    'header_bold': header_bold_ok,
                })

                if not dim_ok:
                    structure_fail_count += 1
                    print(f"FAIL: Table {t_idx} ({label}) — expected {exp_rows}x{exp_cols}, got {n_rows}x{n_cols}")
                else:
                    print(f"  OK: Table {t_idx} ({label}) — {n_rows}x{n_cols} dims correct")

                if not header_bold_ok:
                    structure_fail_count += 1
                    print(f"FAIL: Table {t_idx} ({label}) — header not bold (bold count={header_bold_count})")
                else:
                    print(f"  OK: Table {t_idx} ({label}) — header row is bold")

            if structure_fail_count == 0:
                print(f"PASS: Component 4 — all tables have correct structure and bold headers (0.25 pts)")
                total_score += 0.25
            else:
                # Partial credit: if at least 2 of 3 tables are fully correct
                fully_ok = sum(1 for c in structure_checks if c['dim_ok'] and c['header_bold'])
                if fully_ok >= 2:
                    print(f"PASS (partial): Component 4 — {fully_ok}/3 tables fully correct (0.15 pts)")
                    total_score += 0.15
                elif fully_ok >= 1:
                    print(f"FAIL: Component 4 — only {fully_ok}/3 tables fully correct")
                else:
                    print(f"FAIL: Component 4 — no tables have correct structure")
        else:
            print(f"FAIL: Component 4 — skipped (less than 3 tables)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test the canonical document path on the VM
if not os.path.exists(DOC_PATH):
    print(f"File not found: {DOC_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(DOC_PATH)
