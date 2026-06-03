"""
Initial Setup: Create a Writer document with inline code references in body text
Task ID: writer_tech_025
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_025'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('Technical Setup Guide', level=0)

    # --- Section 1: Getting Started (contains the three inline code references) ---
    doc.add_heading('1. Getting Started', level=1)

    # Paragraph 1 with 'config.yaml'
    p1 = doc.add_paragraph()
    r1a = p1.add_run('Before deploying the application, you must first edit the ')
    r1a.font.size = Pt(11)
    r1a.font.name = 'Liberation Sans'
    r1b = p1.add_run('config.yaml')
    r1b.font.size = Pt(11)
    r1b.font.name = 'Liberation Sans'
    r1c = p1.add_run(' file located in the project root directory. This file contains all the '
                      'environment-specific settings including database connection strings, '
                      'API endpoints, and logging configurations.')
    r1c.font.size = Pt(11)
    r1c.font.name = 'Liberation Sans'

    # Paragraph 2 with 'init()'
    p2 = doc.add_paragraph()
    r2a = p2.add_run('The bootstrap sequence begins when the application calls ')
    r2a.font.size = Pt(11)
    r2a.font.name = 'Liberation Sans'
    r2b = p2.add_run('init()')
    r2b.font.size = Pt(11)
    r2b.font.name = 'Liberation Sans'
    r2c = p2.add_run(' during startup. This function reads the configuration, validates all '
                      'required fields, establishes database connections, and registers '
                      'service handlers with the event dispatcher.')
    r2c.font.size = Pt(11)
    r2c.font.name = 'Liberation Sans'

    # Paragraph 3 with 'ENV_PATH'
    p3 = doc.add_paragraph()
    r3a = p3.add_run('If you need to override the default configuration location, set the ')
    r3a.font.size = Pt(11)
    r3a.font.name = 'Liberation Sans'
    r3b = p3.add_run('ENV_PATH')
    r3b.font.size = Pt(11)
    r3b.font.name = 'Liberation Sans'
    r3c = p3.add_run(' environment variable to point to your custom settings directory. '
                      'The application will search this path first before falling back to '
                      'the default location under /etc/appconfig/.')
    r3c.font.size = Pt(11)
    r3c.font.name = 'Liberation Sans'

    # --- Section 2: Architecture Overview ---
    doc.add_heading('2. Architecture Overview', level=1)

    p4 = doc.add_paragraph()
    r4 = p4.add_run('The system follows a microservices architecture with three primary '
                     'components: the API gateway, the processing engine, and the data '
                     'persistence layer. Each component communicates through a message '
                     'broker using asynchronous event-driven patterns.')
    r4.font.size = Pt(11)
    r4.font.name = 'Liberation Sans'

    p5 = doc.add_paragraph()
    r5 = p5.add_run('The API gateway handles all incoming HTTP requests, performs '
                     'authentication and rate limiting, then routes validated requests '
                     'to the appropriate processing engine instance. Load balancing is '
                     'achieved through consistent hashing of the request payload.')
    r5.font.size = Pt(11)
    r5.font.name = 'Liberation Sans'

    # --- Section 3: Deployment ---
    doc.add_heading('3. Deployment Checklist', level=1)

    items = [
        'Verify all configuration files are present and valid',
        'Run the database migration scripts in sequential order',
        'Start the message broker and confirm cluster health',
        'Deploy the API gateway with the latest TLS certificates',
        'Launch processing engine instances across all availability zones',
        'Execute the integration test suite and verify all endpoints respond',
    ]
    for item in items:
        bp = doc.add_paragraph(style='List Bullet')
        r = bp.add_run(item)
        r.font.size = Pt(11)
        r.font.name = 'Liberation Sans'

    # --- Section 4: Troubleshooting ---
    doc.add_heading('4. Troubleshooting', level=1)

    p6 = doc.add_paragraph()
    r6 = p6.add_run('If the application fails to start, check the system logs located '
                     'in /var/log/appservice/ for detailed error messages. Common issues '
                     'include incorrect file permissions on the configuration directory, '
                     'expired database credentials, and network timeouts when connecting '
                     'to external dependency services.')
    r6.font.size = Pt(11)
    r6.font.name = 'Liberation Sans'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
