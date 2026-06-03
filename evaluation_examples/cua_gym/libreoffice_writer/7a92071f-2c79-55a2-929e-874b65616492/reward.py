"""
Reward Script: Job Description Document for Senior Data Engineer at DataFlow Inc.
Task ID: writer_wf_030
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15): Job title "Senior Data Engineer" as Heading 1, bold
  Component 2 (0.10): Company overview paragraph mentioning "DataFlow Inc."
  Component 3 (0.20): Info table with 4 rows (Department, Location, Reports To, Employment Type)
  Component 4 (0.15): Key Responsibilities section with 6 bullet items
  Component 5 (0.15): Required Qualifications section with 5 bullet items
  Component 6 (0.10): Preferred Qualifications section with 3 bullet items
  Component 7 (0.10): Benefits section with 4 bullet items
  Component 8 (0.05): Application Process section present
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_030'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather all paragraphs and their styles for reuse
    paragraphs = doc.paragraphs
    tables = doc.tables

    if len(paragraphs) == 0 and len(tables) == 0:
        print("FAIL: Document is empty (no paragraphs, no tables)")
        print("REWARD: 0.0")
        return 0.0

    # Helper: collect sections (Heading 2) and their bullet items
    def get_sections_with_bullets():
        """Returns dict: section_name -> list of bullet texts"""
        sections = {}
        current_section = None
        for p in paragraphs:
            style_name = p.style.name if p.style else ''
            if 'Heading 2' in style_name or 'Heading2' in style_name:
                current_section = p.text.strip()
                sections[current_section] = []
            elif current_section and ('List Bullet' in style_name or 'List' in style_name):
                if p.text.strip():
                    sections[current_section].append(p.text.strip())
        return sections

    # Component 1: Job title "Senior Data Engineer" as Heading 1, bold (0.15 points)
    try:
        title_found = False
        for p in paragraphs:
            style_name = p.style.name if p.style else ''
            if ('Heading 1' in style_name or 'Heading1' in style_name):
                text = p.text.strip().lower()
                if 'senior data engineer' in text:
                    # Check if any run is bold
                    has_bold = any(r.bold for r in p.runs if r.text.strip())
                    if has_bold:
                        title_found = True
                        print(f"PASS: Component 1 — Job title 'Senior Data Engineer' found as Heading 1 with bold (0.15 pts)")
                        total_score += 0.15
                    else:
                        # Heading 1 style itself may be bold by default, accept it
                        title_found = True
                        print(f"PASS: Component 1 — Job title 'Senior Data Engineer' found as Heading 1 (0.15 pts)")
                        total_score += 0.15
                    break
        if not title_found:
            print(f"FAIL: Component 1 — No Heading 1 paragraph with 'Senior Data Engineer' found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Company overview paragraph mentioning "DataFlow Inc." (0.10 points)
    try:
        overview_found = False
        for p in paragraphs:
            text = p.text.strip().lower()
            if 'dataflow' in text and len(p.text.strip()) > 50:
                # This should be a regular paragraph (Normal style), not a heading or bullet
                style_name = p.style.name if p.style else ''
                if 'Heading' not in style_name and 'List' not in style_name:
                    overview_found = True
                    print(f"PASS: Component 2 — Company overview paragraph found mentioning DataFlow (0.10 pts)")
                    total_score += 0.10
                    break
        if not overview_found:
            print(f"FAIL: Component 2 — No company overview paragraph mentioning 'DataFlow' found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Info table with 4 rows: Department, Location, Reports To, Employment Type (0.20 points)
    try:
        table_score = 0.0
        if len(tables) >= 1:
            t = tables[0]
            # Check it has at least 4 rows and 2 cols
            if len(t.rows) >= 4 and len(t.columns) >= 2:
                # Collect first-column labels
                labels = [row.cells[0].text.strip().lower() for row in t.rows]
                expected_labels = ['department', 'location', 'reports to', 'employment type']
                found_count = 0
                for exp in expected_labels:
                    if any(exp in label for label in labels):
                        found_count += 1

                if found_count >= 4:
                    table_score = 0.20
                    print(f"PASS: Component 3 — Info table with all 4 required fields (0.20 pts)")
                elif found_count >= 2:
                    table_score = 0.10
                    print(f"PARTIAL: Component 3 — Info table with {found_count}/4 required fields (0.10 pts)")
                else:
                    print(f"FAIL: Component 3 — Info table has {found_count}/4 required fields")
            else:
                print(f"FAIL: Component 3 — Table dimensions too small: {len(t.rows)} rows x {len(t.columns)} cols")
        else:
            print(f"FAIL: Component 3 — No tables found in document")
        total_score += table_score
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Get sections for components 4-8
    sections = get_sections_with_bullets()
    section_names_lower = {k.lower(): (k, v) for k, v in sections.items()}

    # Component 4: Key Responsibilities section with 6 bullet items (0.15 points)
    try:
        resp_key = None
        for k in section_names_lower:
            if 'responsibilit' in k:
                resp_key = k
                break
        if resp_key:
            name, bullets = section_names_lower[resp_key]
            count = len(bullets)
            if count >= 6:
                print(f"PASS: Component 4 — Key Responsibilities has {count} bullet items (0.15 pts)")
                total_score += 0.15
            elif count >= 3:
                partial = round(0.15 * count / 6, 2)
                print(f"PARTIAL: Component 4 — Key Responsibilities has {count}/6 items ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — Key Responsibilities has only {count}/6 items")
        else:
            print(f"FAIL: Component 4 — No 'Key Responsibilities' section found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Required Qualifications section with 5 bullet items (0.15 points)
    try:
        req_key = None
        for k in section_names_lower:
            if 'required' in k and 'qualif' in k:
                req_key = k
                break
        if req_key:
            name, bullets = section_names_lower[req_key]
            count = len(bullets)
            if count >= 5:
                print(f"PASS: Component 5 — Required Qualifications has {count} bullet items (0.15 pts)")
                total_score += 0.15
            elif count >= 3:
                partial = round(0.15 * count / 5, 2)
                print(f"PARTIAL: Component 5 — Required Qualifications has {count}/5 items ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 — Required Qualifications has only {count}/5 items")
        else:
            print(f"FAIL: Component 5 — No 'Required Qualifications' section found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Preferred Qualifications section with 3 bullet items (0.10 points)
    try:
        pref_key = None
        for k in section_names_lower:
            if 'preferred' in k and 'qualif' in k:
                pref_key = k
                break
        if pref_key:
            name, bullets = section_names_lower[pref_key]
            count = len(bullets)
            if count >= 3:
                print(f"PASS: Component 6 — Preferred Qualifications has {count} bullet items (0.10 pts)")
                total_score += 0.10
            elif count >= 1:
                partial = round(0.10 * count / 3, 2)
                print(f"PARTIAL: Component 6 — Preferred Qualifications has {count}/3 items ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 6 — Preferred Qualifications has only {count}/3 items")
        else:
            print(f"FAIL: Component 6 — No 'Preferred Qualifications' section found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Benefits section with 4 bullet items (0.10 points)
    try:
        ben_key = None
        for k in section_names_lower:
            if 'benefit' in k:
                ben_key = k
                break
        if ben_key:
            name, bullets = section_names_lower[ben_key]
            count = len(bullets)
            if count >= 4:
                print(f"PASS: Component 7 — Benefits has {count} bullet items (0.10 pts)")
                total_score += 0.10
            elif count >= 2:
                partial = round(0.10 * count / 4, 2)
                print(f"PARTIAL: Component 7 — Benefits has {count}/4 items ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 7 — Benefits has only {count}/4 items")
        else:
            print(f"FAIL: Component 7 — No 'Benefits' section found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # Component 8: Application Process section present (0.05 points)
    try:
        app_key = None
        for k in section_names_lower:
            if 'application' in k:
                app_key = k
                break
        if app_key:
            print(f"PASS: Component 8 — Application Process section found (0.05 pts)")
            total_score += 0.05
        else:
            # Also check for a Heading 2 with "Application" even if no bullets
            app_heading_found = False
            for p in paragraphs:
                style_name = p.style.name if p.style else ''
                if ('Heading 2' in style_name or 'Heading2' in style_name):
                    if 'application' in p.text.strip().lower():
                        app_heading_found = True
                        break
            if app_heading_found:
                print(f"PASS: Component 8 — Application Process heading found (0.05 pts)")
                total_score += 0.05
            else:
                print(f"FAIL: Component 8 — No 'Application Process' section found")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
