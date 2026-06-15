"""
Reward Script: Create 'Exhibit Page' custom page style with landscape, 0.75in margins, EXHIBIT A header
Task ID: writer_legal_043
Domain: libreoffice_writer
Scoring:
  - Component 1: Document has 2+ sections (0.15)
  - Component 2: Second section is landscape (0.25)
  - Component 3: Second section has 0.75in margins all sides (0.25)
  - Component 4: Second section header contains 'EXHIBIT A' (0.20)
  - Component 5: Second section header is independent (not linked to previous) (0.15)
"""

import os
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_043'

# Tolerance for margin comparison (allow ~2% deviation)
MARGIN_TOLERANCE = 15000  # EMU tolerance (~0.016 inches)
TARGET_MARGIN = Inches(0.75)  # 685800 EMU


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Document has at least 2 sections (0.15 points)
    # Initial has 1 section; golden should have 2+ (section break inserted for Exhibit page)
    try:
        num_sections = len(doc.sections)
        if num_sections >= 2:
            print(f"PASS: Component 1 — Document has {num_sections} sections (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Expected >= 2 sections, found {num_sections}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Gate: if only 1 section, remaining checks cannot pass
    if len(doc.sections) < 2:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Use the last section (the Exhibit section) for remaining checks
    exhibit_section = doc.sections[-1]

    # Component 2: Exhibit section has landscape orientation (0.25 points)
    # Initial is all portrait; golden has landscape on the exhibit section
    try:
        is_landscape = exhibit_section.orientation == WD_ORIENT.LANDSCAPE
        # Also verify width > height (landscape dimensions)
        dims_landscape = exhibit_section.page_width > exhibit_section.page_height
        if is_landscape and dims_landscape:
            print(f"PASS: Component 2 — Landscape orientation confirmed (w={exhibit_section.page_width}, h={exhibit_section.page_height}) (0.25 pts)")
            total_score += 0.25
        elif is_landscape or dims_landscape:
            print(f"PARTIAL: Component 2 — Orientation flag={is_landscape}, dims_landscape={dims_landscape} (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 — Expected landscape, found portrait (orient={exhibit_section.orientation}, w={exhibit_section.page_width}, h={exhibit_section.page_height})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All four margins are 0.75 inches (685800 EMU) (0.25 points)
    # Initial has 1-inch margins; golden should have 0.75-inch margins
    try:
        margins = {
            'left': exhibit_section.left_margin,
            'right': exhibit_section.right_margin,
            'top': exhibit_section.top_margin,
            'bottom': exhibit_section.bottom_margin,
        }
        matching_margins = 0
        for name, val in margins.items():
            if val is not None and abs(val - TARGET_MARGIN) <= MARGIN_TOLERANCE:
                matching_margins += 1
                print(f"  Margin {name}: {val} EMU (target {TARGET_MARGIN}) — OK")
            else:
                print(f"  Margin {name}: {val} EMU (target {TARGET_MARGIN}) — MISMATCH")

        if matching_margins == 4:
            print(f"PASS: Component 3 — All 4 margins are 0.75 inches (0.25 pts)")
            total_score += 0.25
        elif matching_margins >= 2:
            partial = round(0.25 * matching_margins / 4, 2)
            print(f"PARTIAL: Component 3 — {matching_margins}/4 margins match ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {matching_margins}/4 margins match 0.75 inches")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Header contains 'EXHIBIT A' text (0.20 points)
    # Initial has empty header; golden should have 'EXHIBIT A'
    try:
        header = exhibit_section.header
        header_text = ''
        if header and header.paragraphs:
            header_text = ' '.join(p.text.strip() for p in header.paragraphs).strip()

        if 'EXHIBIT A' in header_text.upper():
            print(f"PASS: Component 4 — Header contains 'EXHIBIT A': {repr(header_text)} (0.20 pts)")
            total_score += 0.20
        elif 'EXHIBIT' in header_text.upper():
            print(f"PARTIAL: Component 4 — Header contains 'EXHIBIT' but not 'EXHIBIT A': {repr(header_text)} (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — Header text is {repr(header_text)}, expected to contain 'EXHIBIT A'")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Header is independent (not linked to previous section) (0.15 points)
    # This ensures the Exhibit section has its own distinct header
    try:
        header = exhibit_section.header
        if not header.is_linked_to_previous:
            print(f"PASS: Component 5 — Header is independent (not linked to previous) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — Header is linked to previous (should be independent for distinct 'EXHIBIT A' header)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
