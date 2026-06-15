"""
Initial Setup: Annual HR Report with raw data in paragraph form
Task ID: writer_hr_069
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_069'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title as heading level 0
    doc.add_heading("Annual_HR_Report_2025", level=0)

    # --- Executive Summary (plain text header, raw paragraph) ---
    p = doc.add_paragraph("Executive Summary")
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph(
        "This report presents the human resources performance data for fiscal year 2025. "
        "Key highlights include a total headcount of 2,847 employees across 10 departments, "
        "an average monthly turnover rate of 3.2%, 412 new hires through recruitment efforts, "
        "an overall training completion rate of 78.4%, and continued progress in diversity initiatives. "
        "The compensation benchmarking analysis indicates our pay structure remains competitive "
        "within the 65th percentile of the industry. Further details and data breakdowns are "
        "provided in the sections below."
    )

    # --- Headcount Analytics (raw paragraph data) ---
    p = doc.add_paragraph("Headcount Analytics")
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph(
        "The organization employs 2,847 staff members distributed across 10 departments. "
        "Engineering has 524 employees, Sales has 389, Marketing has 215, Finance has 178, "
        "Human Resources has 142, Operations has 436, Customer Support has 312, "
        "Legal has 87, Research and Development has 298, and IT Infrastructure has 266. "
        "Full-time employees account for 2,314, part-time for 347, and contractors for 186. "
        "Headcount grew by 8.3% year-over-year, with Engineering and R&D seeing the largest increases."
    )

    # --- Turnover Analysis (raw paragraph data, monthly) ---
    p = doc.add_paragraph("Turnover Analysis")
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph(
        "Monthly turnover rates for 2025 were as follows: January 2.8%, February 3.1%, "
        "March 3.4%, April 2.9%, May 3.6%, June 4.2%, July 3.8%, August 3.5%, "
        "September 3.0%, October 2.7%, November 2.6%, December 2.9%. "
        "Voluntary departures accounted for 68% of total separations, with involuntary "
        "terminations at 22% and retirements at 10%. The highest turnover was observed "
        "in Customer Support (5.1% average) and Sales (4.7% average). "
        "Exit interview data indicates that career development opportunities and "
        "compensation were the top two reasons cited by departing employees."
    )

    # --- Recruitment Metrics (raw paragraph data) ---
    p = doc.add_paragraph("Recruitment Metrics")
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph(
        "A total of 412 positions were filled during the fiscal year. The average time-to-fill "
        "was 34 days and cost-per-hire averaged $4,230. Sourcing channels performed as follows: "
        "Employee referrals produced 124 hires with an average time-to-fill of 22 days, "
        "job boards yielded 98 hires at 38 days, LinkedIn sourcing brought in 87 hires at 31 days, "
        "recruitment agencies filled 56 positions at 42 days, campus recruiting added 32 hires at 28 days, "
        "and internal transfers accounted for 15 moves at 14 days. "
        "Offer acceptance rate was 82.6% and first-year retention of new hires stood at 88.1%."
    )

    # --- Training Completion Rates (raw paragraph data) ---
    p = doc.add_paragraph("Training Completion Rates")
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph(
        "Overall training completion rate reached 78.4%. By department: Engineering 85.2%, "
        "Sales 72.1%, Marketing 81.0%, Finance 90.3%, Human Resources 94.7%, Operations 68.5%, "
        "Customer Support 74.8%, Legal 88.9%, Research and Development 82.4%, "
        "and IT Infrastructure 76.3%. Mandatory compliance training achieved 96.2% completion, "
        "leadership development programs 64.8%, and technical skills workshops 71.5%. "
        "Average training hours per employee were 32.6 for the year."
    )

    # --- Compensation Benchmarking (raw paragraph data) ---
    p = doc.add_paragraph("Compensation Benchmarking")
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph(
        "Median salaries by department: Engineering $112,400, Sales $78,600, Marketing $74,200, "
        "Finance $89,500, Human Resources $71,800, Operations $65,300, Customer Support $52,100, "
        "Legal $105,700, Research and Development $108,900, IT Infrastructure $92,600. "
        "Our overall compensation sits at the 65th percentile of the industry benchmark. "
        "Benefits cost per employee averaged $18,400 annually. Variable compensation "
        "(bonuses and commissions) represented 14.2% of total compensation spend."
    )

    # --- Diversity Statistics (raw paragraph data) ---
    p = doc.add_paragraph("Diversity Statistics")
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph(
        "Gender distribution: 54.3% male, 43.8% female, 1.9% non-binary. "
        "At the leadership level (Director and above): 48.2% male, 49.1% female, 2.7% non-binary. "
        "Ethnic diversity breakdown: White 47.2%, Asian 22.8%, Hispanic/Latino 15.4%, "
        "Black/African American 10.1%, Multiracial 3.2%, Other 1.3%. "
        "Age distribution: Under 30 accounts for 28.4%, 30-39 for 34.7%, 40-49 for 22.1%, "
        "50-59 for 11.6%, and 60+ for 3.2%. Employees with disabilities represent 6.8% of the workforce."
    )

    # --- Recommendations (raw paragraph data) ---
    p = doc.add_paragraph("Recommendations")
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph(
        "Based on the analysis presented in this report, the following actions are recommended: "
        "Expand career development programs to address the primary driver of voluntary turnover. "
        "Increase recruitment budget for employee referral bonuses given their superior time-to-fill "
        "and retention metrics. Implement targeted training interventions for Operations and "
        "Customer Support departments where completion rates lag. Conduct a comprehensive "
        "compensation review for Sales and Customer Support roles where turnover is highest. "
        "Continue diversity hiring initiatives with a focus on leadership pipeline development. "
        "Invest in manager training to improve retention and engagement scores."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
