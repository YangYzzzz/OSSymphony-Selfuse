"""
Reward Script: Annual HR Report Document Design
Task ID: writer_hr_069
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Section headings use Heading styles (not plain Normal)
  Component 2 (0.30): At least 6 data tables present with correct structure
  Component 3 (0.15): Table captions present (e.g. "Table 1: ...")
  Component 4 (0.15): Figure placeholders present (e.g. "[Figure X: ...]")
  Component 5 (0.10): In-text table references (e.g. "see Table 1")
  Component 6 (0.10): Recommendations section uses list formatting
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_069'


def persist_app_state(domain: str):
    """Attempt to save any unsaved GUI edits."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            import time
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Section headings use Heading styles (0.20 points)
    # The golden doc converts plain Normal section headers to Heading 1.
    # Expected headings: Executive Summary, Headcount Analytics, Turnover Analysis,
    # Recruitment Metrics, Training Completion Rates, Compensation Benchmarking,
    # Diversity Statistics, Recommendations
    try:
        heading_paras = [
            p for p in doc.paragraphs
            if p.style and p.style.name.startswith('Heading')
        ]
        heading_count = len(heading_paras)
        if heading_count >= 7:
            print(f"PASS: Component 1 — {heading_count} Heading-styled paragraphs found (0.20 pts)")
            total_score += 0.20
        elif heading_count >= 4:
            partial = 0.10
            print(f"PARTIAL: Component 1 — {heading_count} Heading-styled paragraphs (partial {partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — Only {heading_count} Heading-styled paragraphs found, expected >= 7")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: At least 6 data tables present (0.30 points)
    # Initial doc has 0 tables; golden has 6 with structured data.
    try:
        table_count = len(doc.tables)
        if table_count >= 6:
            # Verify tables have reasonable structure (at least 2 rows and 3 columns each)
            well_formed = 0
            for t in doc.tables:
                if len(t.rows) >= 2 and len(t.columns) >= 3:
                    well_formed += 1
            if well_formed >= 6:
                print(f"PASS: Component 2 — {table_count} tables, {well_formed} well-formed (0.30 pts)")
                total_score += 0.30
            elif well_formed >= 3:
                partial = 0.15
                print(f"PARTIAL: Component 2 — {well_formed}/{table_count} well-formed tables (partial {partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — Only {well_formed} well-formed tables out of {table_count}")
        elif table_count >= 3:
            partial = 0.15
            print(f"PARTIAL: Component 2 — {table_count} tables found, expected >= 6 (partial {partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {table_count} tables found, expected >= 6")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Table captions present (0.15 points)
    # Golden doc has paragraphs like "Table 1: Department-wise Headcount Breakdown"
    try:
        caption_pattern = re.compile(r'^Table\s+\d+\s*:', re.IGNORECASE)
        captions = [
            p.text.strip() for p in doc.paragraphs
            if caption_pattern.match(p.text.strip())
        ]
        caption_count = len(captions)
        if caption_count >= 6:
            print(f"PASS: Component 3 — {caption_count} table captions found (0.15 pts)")
            total_score += 0.15
        elif caption_count >= 3:
            partial = 0.08
            print(f"PARTIAL: Component 3 — {caption_count} table captions found (partial {partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {caption_count} table captions found, expected >= 6")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Figure/chart placeholders present (0.15 points)
    # Golden doc has paragraphs like "[Figure 1: Department Headcount Distribution Chart]"
    try:
        figure_pattern = re.compile(r'\[Figure\s+\d+', re.IGNORECASE)
        figure_placeholders = [
            p.text.strip() for p in doc.paragraphs
            if figure_pattern.search(p.text.strip())
        ]
        fig_count = len(figure_placeholders)
        if fig_count >= 6:
            print(f"PASS: Component 4 — {fig_count} figure placeholders found (0.15 pts)")
            total_score += 0.15
        elif fig_count >= 3:
            partial = 0.08
            print(f"PARTIAL: Component 4 — {fig_count} figure placeholders found (partial {partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {fig_count} figure placeholders found, expected >= 6")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: In-text table references (0.10 points)
    # Golden doc has references like "see Table 1", "Table 2", etc. within body text
    # (not counting captions themselves, counting references in descriptive paragraphs)
    try:
        ref_pattern = re.compile(r'(?:see\s+)?Table\s+\d+', re.IGNORECASE)
        # Count paragraphs that contain table references but are NOT themselves captions
        caption_pattern2 = re.compile(r'^Table\s+\d+\s*:', re.IGNORECASE)
        body_paras_with_refs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            if caption_pattern2.match(text):
                continue  # skip captions
            if ref_pattern.search(text):
                body_paras_with_refs.append(text[:80])

        ref_count = len(body_paras_with_refs)
        if ref_count >= 4:
            print(f"PASS: Component 5 — {ref_count} paragraphs with in-text table references (0.10 pts)")
            total_score += 0.10
        elif ref_count >= 2:
            partial = 0.05
            print(f"PARTIAL: Component 5 — {ref_count} paragraphs with table references (partial {partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Only {ref_count} paragraphs with in-text table references, expected >= 4")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Recommendations section uses list formatting (0.10 points)
    # Golden doc has List Bullet style paragraphs in the recommendations section
    try:
        list_paras = [
            p for p in doc.paragraphs
            if p.style and 'List' in p.style.name and p.text.strip()
        ]
        list_count = len(list_paras)
        if list_count >= 3:
            print(f"PASS: Component 6 — {list_count} list-styled paragraphs found in recommendations (0.10 pts)")
            total_score += 0.10
        elif list_count >= 1:
            partial = 0.05
            print(f"PARTIAL: Component 6 — {list_count} list-styled paragraphs (partial {partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 6 — No list-styled paragraphs found for recommendations")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
