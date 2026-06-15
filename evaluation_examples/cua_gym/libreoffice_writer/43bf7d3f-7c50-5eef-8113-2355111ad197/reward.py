"""
Reward Script: Scientific journal article draft in LibreOffice Writer
Task ID: writer_wf_086
Domain: libreoffice_writer
Scoring:
  C1: Title present (0.10)
  C2: 3 authors with affiliations (0.10)
  C3: Abstract section with italic text (0.10)
  C4: Keywords line present (0.05)
  C5: 8 Heading 1 sections with correct names (0.15)
  C6: 3 Heading 2 subsections (Materials, Preparation, Characterization) (0.10)
  C7: Results table with 4 cols and 5+ data rows (0.15)
  C8: 6 references (0.10)
  C9: Double line spacing (0.10)
  C10: Times New Roman body font (0.05)
"""

import os
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_086'


def persist_app_state(domain):
    """Save any unsaved changes in LibreOffice Writer."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    paragraphs = doc.paragraphs
    para_texts = [p.text.strip() for p in paragraphs]

    # Precondition: document must have content
    if len(paragraphs) < 5:
        print(f"FAIL: Document has too few paragraphs ({len(paragraphs)}), appears empty or near-empty")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title present (0.10 points)
    try:
        title_found = any("novel approaches to biodegradable packaging" in t.lower() for t in para_texts)
        if title_found:
            print(f"PASS: Component 1 — Title 'Novel Approaches to Biodegradable Packaging Materials' found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — Title not found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 3 authors with affiliations (0.10 points)
    try:
        # Look for author-like paragraph: multiple names separated by commas
        author_para_found = False
        affiliation_count = 0
        for p in paragraphs:
            text = p.text.strip()
            # Author line: contains multiple names (look for comma-separated names)
            if not author_para_found and ',' in text and len(text) > 10:
                # Check if it looks like author names (not a heading, not too long for affiliations)
                words = text.split(',')
                if len(words) >= 2 and len(text) < 200:
                    # Check for superscript numbers or just multiple names
                    name_indicators = sum(1 for w in words if len(w.strip()) > 3 and any(c.isalpha() for c in w))
                    if name_indicators >= 2:
                        author_para_found = True
            # Affiliation lines: contain university/department/school
            if any(kw in text.lower() for kw in ['university', 'department', 'school', 'institute', 'laboratory']):
                affiliation_count += 1

        if author_para_found and affiliation_count >= 3:
            print(f"PASS: Component 2 — Authors found with {affiliation_count} affiliations (0.10 pts)")
            total_score += 0.10
        elif author_para_found and affiliation_count >= 1:
            print(f"PARTIAL: Component 2 — Authors found but only {affiliation_count} affiliations (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 2 — Authors or affiliations not found (author_para={author_para_found}, affiliations={affiliation_count})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Abstract section with italic text (0.10 points)
    try:
        abstract_heading_idx = None
        for i, p in enumerate(paragraphs):
            if p.style and p.style.name.startswith('Heading') and 'abstract' in p.text.lower():
                abstract_heading_idx = i
                break

        abstract_italic = False
        if abstract_heading_idx is not None and abstract_heading_idx + 1 < len(paragraphs):
            abs_para = paragraphs[abstract_heading_idx + 1]
            abs_text = abs_para.text.strip()
            # Check if the abstract paragraph has substantial text and is italic
            if len(abs_text) > 100:
                italic_runs = [r for r in abs_para.runs if r.font.italic]
                if italic_runs:
                    abstract_italic = True

        if abstract_heading_idx is not None and abstract_italic:
            print(f"PASS: Component 3 — Abstract heading found with italic text (0.10 pts)")
            total_score += 0.10
        elif abstract_heading_idx is not None:
            print(f"PARTIAL: Component 3 — Abstract heading found but text not italic (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 3 — Abstract section not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Keywords line present (0.05 points)
    try:
        keywords_found = any('keywords' in t.lower() and len(t) > 15 for t in para_texts)
        if keywords_found:
            print(f"PASS: Component 4 — Keywords line found (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 4 — Keywords line not found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: 8 Heading 1 sections with correct names (0.15 points)
    try:
        expected_h1 = ['abstract', 'introduction', 'materials and methods', 'results',
                       'discussion', 'conclusion', 'acknowledgments', 'references']
        actual_h1 = []
        for p in paragraphs:
            if p.style and p.style.name == 'Heading 1':
                actual_h1.append(p.text.strip().lower())

        matched = sum(1 for exp in expected_h1 if any(exp in act for act in actual_h1))
        if matched >= 8:
            print(f"PASS: Component 5 — All 8 Heading 1 sections found: {[p.text for p in paragraphs if p.style and p.style.name == 'Heading 1']} (0.15 pts)")
            total_score += 0.15
        elif matched >= 5:
            pts = round(0.15 * matched / 8, 2)
            print(f"PARTIAL: Component 5 — {matched}/8 Heading 1 sections found ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 5 — Only {matched}/8 Heading 1 sections found. Actual H1: {actual_h1}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: 3 Heading 2 subsections (Materials, Preparation, Characterization) (0.10 points)
    try:
        expected_h2 = ['materials', 'preparation', 'characterization']
        actual_h2 = []
        for p in paragraphs:
            if p.style and p.style.name == 'Heading 2':
                actual_h2.append(p.text.strip().lower())

        matched_h2 = sum(1 for exp in expected_h2 if any(exp in act for act in actual_h2))
        if matched_h2 >= 3:
            print(f"PASS: Component 6 — All 3 Heading 2 subsections found: {actual_h2} (0.10 pts)")
            total_score += 0.10
        elif matched_h2 >= 1:
            pts = round(0.10 * matched_h2 / 3, 2)
            print(f"PARTIAL: Component 6 — {matched_h2}/3 Heading 2 subsections found ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 6 — No Heading 2 subsections found. Actual H2: {actual_h2}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Results table with correct structure (0.15 points)
    try:
        tables = doc.tables
        table_found = False
        correct_headers = False
        correct_rows = False

        if len(tables) >= 1:
            table_found = True
            # Find the results table (should have Sample, Tensile Strength, etc.)
            for tbl in tables:
                num_rows = len(tbl.rows)
                num_cols = len(tbl.columns)
                if num_cols >= 4 and num_rows >= 6:  # header + 5 data rows
                    header_cells = [c.text.strip().lower() for c in tbl.rows[0].cells]
                    has_sample = any('sample' in h for h in header_cells)
                    has_tensile = any('tensile' in h for h in header_cells)
                    has_elongation = any('elongation' in h for h in header_cells)
                    has_degradation = any('degradation' in h or 'degrad' in h for h in header_cells)

                    if has_sample and has_tensile:
                        correct_headers = True
                    if num_rows >= 6:  # 1 header + 5 data
                        correct_rows = True
                    break

        if table_found and correct_headers and correct_rows:
            print(f"PASS: Component 7 — Results table with correct headers and 5+ data rows (0.15 pts)")
            total_score += 0.15
        elif table_found and (correct_headers or correct_rows):
            print(f"PARTIAL: Component 7 — Table found but headers={correct_headers}, rows={correct_rows} (0.08 pts)")
            total_score += 0.08
        elif table_found:
            print(f"PARTIAL: Component 7 — Table found but structure incorrect (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 7 — No table found in document")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # Component 8: 6 references (0.10 points)
    try:
        # Find References section and count entries after it
        ref_heading_idx = None
        for i, p in enumerate(paragraphs):
            if p.style and p.style.name == 'Heading 1' and 'reference' in p.text.lower():
                ref_heading_idx = i
                break

        ref_count = 0
        if ref_heading_idx is not None:
            for i in range(ref_heading_idx + 1, len(paragraphs)):
                p = paragraphs[i]
                # Stop if we hit another heading
                if p.style and p.style.name.startswith('Heading'):
                    break
                text = p.text.strip()
                # Count non-empty paragraphs as reference entries
                if len(text) > 20:
                    ref_count += 1

        if ref_count >= 6:
            print(f"PASS: Component 8 — {ref_count} references found (0.10 pts)")
            total_score += 0.10
        elif ref_count >= 3:
            pts = round(0.10 * ref_count / 6, 2)
            print(f"PARTIAL: Component 8 — {ref_count}/6 references found ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 8 — Only {ref_count} references found")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    # Component 9: Double line spacing (0.10 points)
    try:
        # Check line spacing on body paragraphs (Normal style)
        double_spaced_count = 0
        body_para_count = 0
        for p in paragraphs:
            if p.style and p.style.name == 'Normal' and p.text.strip():
                body_para_count += 1
                ls = p.paragraph_format.line_spacing
                if ls is not None and abs(float(ls) - 2.0) < 0.1:
                    double_spaced_count += 1

        if body_para_count > 0 and double_spaced_count >= body_para_count * 0.8:
            print(f"PASS: Component 9 — Double spacing on {double_spaced_count}/{body_para_count} body paragraphs (0.10 pts)")
            total_score += 0.10
        elif body_para_count > 0 and double_spaced_count >= body_para_count * 0.5:
            print(f"PARTIAL: Component 9 — Double spacing on {double_spaced_count}/{body_para_count} body paragraphs (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 9 — Double spacing found on {double_spaced_count}/{body_para_count} body paragraphs")
    except Exception as e:
        print(f"ERROR: Component 9 — {e}")

    # Component 10: Times New Roman body font (0.05 points)
    try:
        tnr_count = 0
        total_runs_with_font = 0
        for p in paragraphs:
            if p.style and p.style.name == 'Normal' and p.text.strip():
                for run in p.runs:
                    if run.font.name:
                        total_runs_with_font += 1
                        if 'times' in run.font.name.lower():
                            tnr_count += 1

        if total_runs_with_font > 0 and tnr_count >= total_runs_with_font * 0.8:
            print(f"PASS: Component 10 — Times New Roman on {tnr_count}/{total_runs_with_font} body runs (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 10 — Times New Roman on {tnr_count}/{total_runs_with_font} body runs")
    except Exception as e:
        print(f"ERROR: Component 10 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
