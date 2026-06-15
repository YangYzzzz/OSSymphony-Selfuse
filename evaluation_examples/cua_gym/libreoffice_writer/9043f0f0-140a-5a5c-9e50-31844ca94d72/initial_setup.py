"""
Initial Setup: Technical reference document with 4 heading levels, no TOC
Task ID: osworld_writer_toc_generation_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_toc_generation_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Document title (not a heading, just a title paragraph)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Enterprise Network Architecture Reference Guide")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(12)

    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run("Technical Documentation — Version 3.2")
    subtitle_run.font.size = Pt(12)
    subtitle_run.italic = True
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_para.paragraph_format.space_after = Pt(24)

    doc.add_paragraph(
        "This reference guide provides comprehensive documentation for the enterprise network "
        "architecture, including hardware specifications, software configurations, security "
        "protocols, and operational procedures. This document is intended for network engineers, "
        "system administrators, and IT managers responsible for maintaining the infrastructure."
    )

    # ── Chapter 1 ──────────────────────────────────────────────────
    doc.add_heading("Network Infrastructure Overview", level=1)
    doc.add_paragraph(
        "The enterprise network infrastructure consists of multiple interconnected components "
        "designed for high availability, security, and performance. This chapter provides an "
        "overview of the primary architectural components and their relationships."
    )

    doc.add_heading("Physical Layer Components", level=2)
    doc.add_paragraph(
        "The physical layer encompasses all hardware elements including switches, routers, "
        "firewalls, and cabling infrastructure. Proper physical layer design is fundamental "
        "to overall network performance and reliability."
    )

    doc.add_heading("Core Switching Fabric", level=3)
    doc.add_paragraph(
        "The core switching fabric utilizes Cisco Catalyst 9500 series switches operating "
        "in a redundant pair configuration. Each switch provides 48 ports of 10GbE connectivity "
        "with uplink capacity of 40GbE per switch pair."
    )

    doc.add_heading("Switch Stack Configuration", level=4)
    doc.add_paragraph(
        "Individual switches are configured in stacks of up to 8 units using StackWise-480 "
        "technology. Each stack operates with a single management plane while providing "
        "480 Gbps of stack interconnect bandwidth."
    )

    doc.add_heading("Port Channel Specifications", level=4)
    doc.add_paragraph(
        "Port channels are configured using LACP (802.3ad) with a minimum of 2 active "
        "links per channel. Maximum bundle size is limited to 8 physical ports with "
        "load balancing configured for source-destination IP hashing."
    )

    doc.add_heading("Distribution Layer Design", level=3)
    doc.add_paragraph(
        "The distribution layer provides routing, policy enforcement, and inter-VLAN "
        "connectivity. Distribution switches connect the access layer to the core and "
        "implement quality of service policies."
    )

    doc.add_heading("VLAN Architecture", level=4)
    doc.add_paragraph(
        "VLANs are allocated in blocks of 100 per department: Management (VLAN 10-19), "
        "Server Farm (VLAN 100-199), User Access (VLAN 200-399), and Guest Network (VLAN 900-999). "
        "Spanning Tree Protocol uses RSTP with distribution switches as primary roots."
    )

    doc.add_heading("Inter-VLAN Routing", level=4)
    doc.add_paragraph(
        "Inter-VLAN routing is performed at the distribution layer using Layer 3 interfaces. "
        "Each VLAN is assigned a /24 subnet with the default gateway residing on the primary "
        "distribution switch. Hot Standby Router Protocol (HSRP) provides gateway redundancy."
    )

    doc.add_heading("Logical Network Architecture", level=2)
    doc.add_paragraph(
        "The logical network architecture defines addressing schemes, routing protocols, "
        "and traffic segmentation policies that govern how data flows through the infrastructure."
    )

    doc.add_heading("IP Addressing Scheme", level=3)
    doc.add_paragraph(
        "The network uses RFC 1918 private address space with a hierarchical allocation "
        "model. The primary address block 10.0.0.0/8 is subdivided by region and function "
        "to facilitate summarization and policy enforcement."
    )

    doc.add_heading("IPv6 Implementation", level=4)
    doc.add_paragraph(
        "IPv6 deployment uses dual-stack configuration on all infrastructure devices. "
        "The globally assigned /32 prefix is sub-delegated to individual sites using /48 "
        "blocks. End-user subnets receive /64 allocations per IEEE recommendations."
    )

    doc.add_heading("DNS Architecture", level=4)
    doc.add_paragraph(
        "DNS services are provided by a pair of BIND 9.16 servers in master-slave "
        "configuration. Internal zones use split-horizon DNS to serve different records "
        "to internal and external clients. DNSSEC is enabled on all authoritative zones."
    )

    doc.add_heading("Routing Protocol Configuration", level=3)
    doc.add_paragraph(
        "OSPF Area 0 serves as the backbone routing domain connecting all core and "
        "distribution switches. Stub areas are configured for branch locations to reduce "
        "LSA flooding and routing table size at remote sites."
    )

    doc.add_heading("OSPF Area Design", level=4)
    doc.add_paragraph(
        "Area 0 (backbone) connects all data center and campus core routers. "
        "Branch areas are configured as NSSA to allow redistribution of external routes "
        "while reducing full LSA flooding. ABR routers perform route summarization at "
        "area boundaries to minimize routing table size."
    )

    doc.add_heading("BGP Peering Policy", level=4)
    doc.add_paragraph(
        "External BGP peering is established with two upstream ISPs for redundancy. "
        "Inbound routing policy uses AS-path prepending and local preference modification "
        "to prefer primary ISP for most traffic. MED attributes are used to influence "
        "inbound traffic from each provider."
    )

    # ── Chapter 2 ──────────────────────────────────────────────────
    doc.add_heading("Security Architecture", level=1)
    doc.add_paragraph(
        "Network security is implemented using a defense-in-depth strategy with multiple "
        "overlapping security controls. This chapter describes the security architecture "
        "including perimeter defenses, internal segmentation, and monitoring systems."
    )

    doc.add_heading("Perimeter Security", level=2)
    doc.add_paragraph(
        "The network perimeter is protected by next-generation firewalls operating in "
        "active-active high availability configuration. All external traffic passes through "
        "dedicated security inspection zones before reaching internal resources."
    )

    doc.add_heading("Firewall Policy Framework", level=3)
    doc.add_paragraph(
        "Firewall policies follow a default-deny model with explicit permits for approved "
        "traffic flows. Rules are organized by traffic classification: Internet-to-DMZ, "
        "Internet-to-Internal, DMZ-to-Internal, and Internal-to-Internet."
    )

    doc.add_heading("Application Layer Inspection", level=4)
    doc.add_paragraph(
        "Deep packet inspection is enabled for HTTP, HTTPS, DNS, and SMTP traffic. "
        "SSL/TLS inspection decrypts and re-encrypts HTTPS traffic for content analysis. "
        "Certificate verification exceptions are documented and reviewed quarterly."
    )

    doc.add_heading("Geo-IP Filtering", level=4)
    doc.add_paragraph(
        "Geographic IP filtering blocks inbound connections from high-risk country codes. "
        "The blocked country list is maintained by the security operations team and updated "
        "monthly based on threat intelligence feeds. Legitimate business exceptions require "
        "documented approval from the CISO."
    )

    doc.add_heading("Intrusion Detection and Prevention", level=3)
    doc.add_paragraph(
        "IDS/IPS sensors are deployed inline at the internet edge and in tap mode at "
        "internal segment boundaries. Snort 3.x rules are updated automatically from the "
        "Cisco Talos intelligence feed with manual review for high-severity signatures."
    )

    doc.add_heading("Signature Management", level=4)
    doc.add_paragraph(
        "Signature updates are applied to staging IPS systems first and monitored for "
        "false positives for 48 hours before production deployment. Critical signatures "
        "addressing active exploitation may bypass the staging requirement with CISO approval."
    )

    doc.add_heading("Alert Tuning Procedures", level=4)
    doc.add_paragraph(
        "IDS alerts are reviewed by SOC analysts daily. False positives are suppressed "
        "using asset-based tuning to reduce alert fatigue. All suppression rules require "
        "dual-analyst review and expire after 90 days unless renewed."
    )

    doc.add_heading("Zero Trust Network Access", level=2)
    doc.add_paragraph(
        "Zero Trust Architecture principles are applied to all user access scenarios. "
        "Users must authenticate with multi-factor credentials before accessing any network "
        "resources regardless of their physical location or network segment."
    )

    doc.add_heading("Identity Provider Integration", level=3)
    doc.add_paragraph(
        "Azure Active Directory serves as the primary identity provider using SAML 2.0 "
        "and OIDC protocols. Conditional Access policies enforce MFA for all cloud and "
        "on-premises application access based on user risk scores and device compliance."
    )

    doc.add_heading("Device Compliance Requirements", level=4)
    doc.add_paragraph(
        "Endpoints must meet minimum compliance requirements: current OS patch level, "
        "active endpoint protection software, encrypted disk storage, and device certificate "
        "enrollment. Non-compliant devices are quarantined to a remediation VLAN with "
        "limited access to update and compliance services only."
    )

    doc.add_heading("Continuous Verification", level=4)
    doc.add_paragraph(
        "Authentication tokens are valid for 4 hours maximum before requiring re-authentication. "
        "Anomalous user behavior triggers step-up authentication challenges. Session monitoring "
        "detects and terminates suspicious connections based on behavioral baselines."
    )

    doc.add_heading("Privileged Access Management", level=3)
    doc.add_paragraph(
        "Administrative access to all systems is brokered through a Privileged Access "
        "Workstation (PAW) infrastructure. Network device management requires jump server "
        "access with session recording and real-time monitoring by the security team."
    )

    doc.add_heading("Just-In-Time Access", level=4)
    doc.add_paragraph(
        "Administrative credentials are provisioned on-demand for approved maintenance "
        "windows. Access requests must include justification and are approved by the "
        "system owner or security operations manager. All sessions are recorded and "
        "retained for 180 days for audit purposes."
    )

    doc.add_heading("Credential Vault Operations", level=4)
    doc.add_paragraph(
        "CyberArk Privilege Cloud manages all privileged credentials. Passwords are "
        "rotated automatically after each use for shared accounts. Service account passwords "
        "rotate on a 30-day schedule. Credential checkouts are logged with user identity, "
        "timestamp, and system accessed."
    )

    # ── Chapter 3 ──────────────────────────────────────────────────
    doc.add_heading("Performance and Capacity Management", level=1)
    doc.add_paragraph(
        "This chapter describes the methodologies and tools used to monitor, measure, "
        "and manage network performance. Capacity planning processes ensure the infrastructure "
        "scales to meet growing business requirements."
    )

    doc.add_heading("Network Monitoring Framework", level=2)
    doc.add_paragraph(
        "Comprehensive monitoring provides visibility into all aspects of network health "
        "and performance. Multiple monitoring systems collect data at different time scales "
        "and granularities to support both real-time operations and long-term trend analysis."
    )

    doc.add_heading("SNMP Polling Infrastructure", level=3)
    doc.add_paragraph(
        "LibreNMS polls all network devices via SNMP v3 every 5 minutes for standard "
        "interface counters. Critical infrastructure polling is performed every 60 seconds "
        "to enable faster fault detection. All devices use unique SNMP v3 credentials stored "
        "in the credential vault."
    )

    doc.add_heading("MIB Extensions", level=4)
    doc.add_paragraph(
        "Vendor-specific MIBs are loaded for Cisco, Juniper, and Palo Alto devices to "
        "expose hardware health metrics including CPU temperature, power supply status, "
        "and fan speed. Custom OIDs are polled for application-specific performance counters."
    )

    doc.add_heading("Threshold Configuration", level=4)
    doc.add_paragraph(
        "Interface utilization thresholds are set at 75% for warning and 90% for critical "
        "alerts. CPU utilization thresholds are 60% warning and 80% critical sustained over "
        "5 minutes. Memory thresholds are 70% warning and 85% critical for network operating systems."
    )

    doc.add_heading("Flow Analysis", level=3)
    doc.add_paragraph(
        "NetFlow v9 data is exported from all core and distribution routers to a centralized "
        "flow collector. Flow data is retained for 90 days and analyzed for traffic patterns, "
        "anomaly detection, and capacity planning purposes."
    )

    doc.add_heading("Top Talkers Analysis", level=4)
    doc.add_paragraph(
        "Weekly automated reports identify top bandwidth consumers by source IP, destination IP, "
        "application protocol, and AS number. Unusual bandwidth consumption patterns trigger "
        "automatic alerts for security and operations team review."
    )

    doc.add_heading("Application Visibility", level=4)
    doc.add_paragraph(
        "NBAR2 application recognition classifies network flows by business application. "
        "Application-level reporting enables capacity planning by workload rather than "
        "pure bandwidth metrics. SaaS application traffic is identified and reported separately."
    )

    doc.add_heading("Capacity Planning Process", level=2)
    doc.add_paragraph(
        "Quarterly capacity reviews assess infrastructure utilization against growth projections. "
        "The capacity planning process incorporates business forecasts, historical growth trends, "
        "and planned project impacts to generate 12-month infrastructure roadmaps."
    )

    doc.add_heading("Bandwidth Forecasting", level=3)
    doc.add_paragraph(
        "Traffic growth projections are based on 18-month historical averages with adjustments "
        "for known business initiatives. WAN circuits are sized to maintain peak utilization "
        "below 60% to preserve headroom for traffic bursts and planned maintenance activities."
    )

    doc.add_heading("Seasonal Adjustment Factors", level=4)
    doc.add_paragraph(
        "Retail sector deployments incorporate seasonal traffic multipliers: Q4 holiday traffic "
        "peaks at 2.4x baseline, summer peaks at 0.7x baseline. Financial sector deployments "
        "apply month-end and quarter-end multipliers of 1.6x and 1.9x respectively."
    )

    doc.add_heading("Growth Rate Calculations", level=4)
    doc.add_paragraph(
        "Compound annual growth rates are calculated for each traffic segment independently. "
        "Video conferencing traffic historically grows at 35% CAGR, while traditional "
        "data replication grows at 8-12% CAGR. Blended growth rates weight each segment "
        "by current traffic volume."
    )

    doc.add_heading("Hardware Lifecycle Management", level=3)
    doc.add_paragraph(
        "Network hardware is replaced on a 5-year lifecycle for core infrastructure and "
        "7-year lifecycle for access layer equipment. End-of-life dates are tracked in the "
        "CMDB with automated alerts generated 18 months before contract end or EOL dates."
    )

    doc.add_heading("Refresh Planning", level=4)
    doc.add_paragraph(
        "Hardware refresh projects are planned 12-18 months in advance to accommodate "
        "procurement lead times and change freeze windows. Refresh projects include "
        "capability upgrades where the new platform provides higher performance, additional "
        "features, or reduced operational complexity compared to replaced equipment."
    )

    doc.add_heading("Decommissioning Procedures", level=4)
    doc.add_paragraph(
        "Equipment decommissioning follows a documented 6-step process: service migration, "
        "traffic verification, physical disconnection, configuration backup archival, "
        "secure data erasure, and asset disposition. All steps require sign-off from "
        "the network operations manager before physical removal."
    )

    # ── Chapter 4 ──────────────────────────────────────────────────
    doc.add_heading("Operational Procedures", level=1)
    doc.add_paragraph(
        "This chapter documents the standard operating procedures for routine network "
        "operations including change management, incident response, and maintenance activities."
    )

    doc.add_heading("Change Management", level=2)
    doc.add_paragraph(
        "All changes to production network infrastructure follow the ITIL-aligned change "
        "management process. Changes are categorized by risk level with appropriate approval "
        "and implementation requirements for each category."
    )

    doc.add_heading("Change Classification Framework", level=3)
    doc.add_paragraph(
        "Changes are classified as Standard (pre-approved, low risk), Normal (requires CAB approval), "
        "or Emergency (expedited approval for service restoration). Emergency changes require "
        "retrospective review within 5 business days of implementation."
    )

    doc.add_heading("Pre-Implementation Checklist", level=4)
    doc.add_paragraph(
        "All Normal changes must complete a pre-implementation checklist: technical review, "
        "rollback plan documented, maintenance window scheduled, stakeholder notifications sent, "
        "and test environment validation completed where applicable. Incomplete checklists "
        "result in automatic CAB rejection."
    )

    doc.add_heading("Post-Implementation Verification", level=4)
    doc.add_paragraph(
        "Changes include verification steps confirmed within 30 minutes of implementation. "
        "Verification includes: configuration audit against change plan, connectivity testing "
        "from monitoring systems, and confirmation from application owners that services "
        "are functioning normally. Failed verification triggers immediate rollback."
    )

    doc.add_heading("Incident Response Procedures", level=3)
    doc.add_paragraph(
        "Network incidents are triaged using a severity matrix based on service impact "
        "and affected user count. Severity 1 incidents require immediate response with "
        "15-minute status updates until resolution."
    )

    doc.add_heading("Escalation Matrix", level=4)
    doc.add_paragraph(
        "Severity 1 (complete outage): NOC engineer + Team Lead notified immediately, "
        "Network Manager escalated at T+15 minutes, Director at T+30 minutes. "
        "Severity 2 (partial degradation): NOC engineer responds within 15 minutes, "
        "Team Lead notified at T+30 minutes if not resolved."
    )

    doc.add_heading("War Room Procedures", level=4)
    doc.add_paragraph(
        "Major incidents (Sev1 exceeding 30 minutes) trigger war room protocols. "
        "A dedicated conference bridge is established with NOC, network engineering, "
        "and affected application teams. A scribe documents all actions taken and their "
        "results. Status updates are published to the incident portal every 30 minutes."
    )

    doc.add_heading("Preventive Maintenance", level=2)
    doc.add_paragraph(
        "Preventive maintenance activities are scheduled during approved maintenance windows "
        "to minimize service impact. Regular maintenance prevents hardware failures and "
        "ensures optimal system performance."
    )

    doc.add_heading("Monthly Maintenance Tasks", level=3)
    doc.add_paragraph(
        "Monthly tasks include: firmware compliance verification, certificate expiration "
        "checks (alert on 90-day threshold), BGP peer health review, unused port audits, "
        "and backup configuration validation. Task completion is tracked in the ITSM system."
    )

    doc.add_heading("Firmware Update Procedures", level=4)
    doc.add_paragraph(
        "Firmware updates are tested in lab environment before production deployment. "
        "Production updates begin with lowest-risk devices (access layer) before core "
        "infrastructure. Update windows are 4 hours with 2-hour rollback time reserved. "
        "Updates affecting more than 20 devices require a dedicated maintenance weekend."
    )

    doc.add_heading("Configuration Backup Validation", level=4)
    doc.add_paragraph(
        "Configuration backups are verified by restoring a random sample of 5 devices "
        "monthly to a staging environment and confirming functionality. Backup age is "
        "monitored with alerts for devices where backup is older than 7 days. Critical "
        "devices require daily backups with weekly integrity verification."
    )

    doc.add_heading("Quarterly Audit Procedures", level=3)
    doc.add_paragraph(
        "Quarterly audits review security configurations, access controls, and compliance "
        "status against internal standards and regulatory requirements. Audit findings are "
        "tracked to remediation with defined timelines based on risk severity."
    )

    doc.add_heading("Access Control Audit", level=4)
    doc.add_paragraph(
        "Quarterly access reviews verify that all active management credentials correspond "
        "to active employees with valid business justification. Terminated employee accounts "
        "must be disabled within 4 hours of HR notification. Service accounts are reviewed "
        "for continued necessity and appropriate privilege levels."
    )

    doc.add_heading("Security Baseline Verification", level=4)
    doc.add_paragraph(
        "All network devices are scanned against the approved security baseline using "
        "automated configuration compliance tools. Deviations are categorized as critical "
        "(must remediate within 72 hours), major (30 days), or minor (90 days). "
        "Remediation status is reported to the CISO monthly."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
