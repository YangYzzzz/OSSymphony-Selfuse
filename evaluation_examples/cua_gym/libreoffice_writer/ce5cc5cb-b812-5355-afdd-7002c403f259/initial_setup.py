"""
Initial Setup: Insert Roman numeral page numbers in footer
Task ID: writer_fs_074
Domain: libreoffice_writer

Creates an 8-page Writer document with footer enabled but empty.
No page number field; numbering format is default Arabic.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_074'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Enable footer but leave it empty
    footer = section.footer
    footer.is_linked_to_previous = False
    # Ensure at least one paragraph exists in footer (empty)
    if not footer.paragraphs:
        footer._element.append(footer._element.makeelement(qn('w:p'), {}))
    # Footer paragraph is empty - no page number field
    fp = footer.paragraphs[0]
    fp.text = ""

    # --- Page 1: Title and Introduction ---
    heading = doc.add_heading("Annual Performance Review Report", level=1)
    doc.add_paragraph("")
    doc.add_paragraph("Prepared by: Human Resources Department")
    doc.add_paragraph("Meridian Technologies Inc.")
    doc.add_paragraph("Date: March 15, 2025")
    doc.add_paragraph("")
    p = doc.add_paragraph(
        "This report provides a comprehensive overview of employee performance "
        "evaluations conducted during the fiscal year 2024-2025. The assessment "
        "covers all departments and includes individual performance metrics, "
        "team contributions, and recommendations for professional development."
    )
    p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        "The evaluation process followed the company's standardized framework, "
        "incorporating 360-degree feedback from supervisors, peers, and direct "
        "reports. Each employee was rated across five core competency areas: "
        "technical proficiency, collaboration, leadership, innovation, and "
        "client engagement."
    )

    # --- Page 2: Engineering Department ---
    doc.add_page_break()
    doc.add_heading("Engineering Department Overview", level=2)
    doc.add_paragraph(
        "The Engineering department demonstrated exceptional performance this year, "
        "completing 94% of planned deliverables ahead of schedule. Lead engineer "
        "Sarah Chen spearheaded the migration to microservices architecture, "
        "reducing system downtime by 37% compared to the previous fiscal year."
    )
    doc.add_paragraph(
        "Notable achievements include the successful launch of the Aurora platform, "
        "which processed over 2.3 million transactions in its first quarter. "
        "The team also implemented automated CI/CD pipelines that reduced "
        "deployment time from 4 hours to 23 minutes."
    )
    doc.add_paragraph(
        "Key personnel highlights:"
    )
    doc.add_paragraph("Sarah Chen - Principal Engineer: Exceeded expectations in all five competency areas. Led the Aurora platform initiative.", style="List Bullet")
    doc.add_paragraph("Marcus Rivera - Senior Developer: Demonstrated outstanding collaboration skills. Mentored three junior developers.", style="List Bullet")
    doc.add_paragraph("Aisha Patel - DevOps Lead: Achieved 99.97% uptime for production systems. Implemented comprehensive monitoring.", style="List Bullet")
    doc.add_paragraph("James Okonkwo - Software Architect: Designed the new event-driven architecture adopted by four product teams.", style="List Bullet")

    # --- Page 3: Marketing Department ---
    doc.add_page_break()
    doc.add_heading("Marketing Department Overview", level=2)
    doc.add_paragraph(
        "The Marketing team exceeded their quarterly targets by 18%, driving "
        "a 23% increase in qualified leads year-over-year. The department's "
        "digital transformation initiative, led by Creative Director Elena "
        "Vasquez, resulted in a 45% improvement in campaign conversion rates."
    )
    doc.add_paragraph(
        "The brand refresh project completed in Q2 received industry recognition, "
        "winning the 2024 Digital Marketing Excellence Award. Social media "
        "engagement increased by 67% following the implementation of the new "
        "content strategy framework developed by the analytics team."
    )
    doc.add_paragraph(
        "Budget utilization was optimized through data-driven allocation, "
        "resulting in a 12% reduction in cost-per-acquisition while maintaining "
        "lead quality standards. The marketing automation platform processed "
        "over 850,000 personalized email campaigns with an average open rate "
        "of 28.4%, well above the industry benchmark of 21.3%."
    )

    # --- Page 4: Sales Department ---
    doc.add_page_break()
    doc.add_heading("Sales Department Overview", level=2)
    doc.add_paragraph(
        "The Sales department achieved $14.7 million in total revenue, "
        "representing a 19% increase over the previous year. Regional "
        "Director Tomoko Ishikawa led the APAC expansion that secured "
        "twelve new enterprise clients valued at $3.2 million collectively."
    )
    doc.add_paragraph(
        "The inside sales team, managed by Derek Fitzgerald, implemented "
        "a consultative selling methodology that increased average deal "
        "size by 34%. Customer retention rate improved to 91.5%, up from "
        "86.2% in the previous period, attributable to the enhanced "
        "account management protocols introduced in Q1."
    )
    doc.add_paragraph(
        "Training initiatives included quarterly skill development workshops "
        "and the introduction of the Sales Excellence certification program. "
        "Fifteen team members completed the advanced negotiation module, and "
        "eight achieved the newly established Platinum Seller designation."
    )

    # --- Page 5: Finance and Operations ---
    doc.add_page_break()
    doc.add_heading("Finance and Operations Review", level=2)
    doc.add_paragraph(
        "The Finance team, under CFO Patricia Morales, delivered outstanding "
        "results in cost management and financial planning. Operating expenses "
        "were reduced by 8.3% through strategic vendor renegotiations and "
        "process automation initiatives."
    )
    doc.add_paragraph(
        "The accounts receivable team reduced days sales outstanding from "
        "42 to 31 days, improving cash flow by approximately $2.1 million. "
        "The annual audit was completed with zero material findings for the "
        "third consecutive year, reflecting the team's commitment to "
        "compliance and accuracy."
    )
    doc.add_paragraph(
        "Operations manager Rafael Santos streamlined warehouse logistics, "
        "achieving a 22% improvement in order fulfillment speed. The "
        "implementation of the new ERP module in Q3 consolidated three "
        "separate inventory systems into a unified platform, eliminating "
        "data reconciliation delays that previously consumed 15 hours "
        "per week."
    )

    # --- Page 6: Human Resources ---
    doc.add_page_break()
    doc.add_heading("Human Resources Initiatives", level=2)
    doc.add_paragraph(
        "The HR department focused on talent acquisition and employee "
        "well-being programs throughout the fiscal year. A total of 47 "
        "new positions were filled with an average time-to-hire of 28 days, "
        "a 15% improvement over the previous year's average of 33 days."
    )
    doc.add_paragraph(
        "The employee wellness program, coordinated by Benefits Manager "
        "Lucia Fernandez, saw a participation rate of 73%, up from 51%. "
        "The program included mental health resources, fitness subsidies, "
        "and flexible work arrangements that contributed to a 4.2-point "
        "improvement in the annual employee satisfaction survey."
    )
    doc.add_paragraph(
        "Diversity and inclusion efforts resulted in a 31% increase in "
        "applications from underrepresented groups. The new mentorship "
        "program paired 24 senior leaders with emerging talent across "
        "departments, with 88% of participants reporting significant "
        "professional growth."
    )

    # --- Page 7: Research and Development ---
    doc.add_page_break()
    doc.add_heading("Research and Development", level=2)
    doc.add_paragraph(
        "The R&D division filed 14 patent applications this fiscal year, "
        "with seven receiving preliminary approval. Dr. Nadia Kowalski's "
        "team achieved a breakthrough in natural language processing that "
        "reduced model training time by 60% while maintaining accuracy "
        "benchmarks."
    )
    doc.add_paragraph(
        "The innovation lab completed 23 proof-of-concept projects, with "
        "five advancing to the production pipeline. The predictive analytics "
        "module developed by the data science team was integrated into "
        "three client-facing products, generating $1.8 million in new "
        "recurring revenue."
    )
    doc.add_paragraph(
        "Research partnerships were established with Stanford University "
        "and the Max Planck Institute, providing access to cutting-edge "
        "facilities and collaborative opportunities. The annual R&D "
        "conference hosted 340 attendees from 12 countries."
    )

    # --- Page 8: Recommendations and Next Steps ---
    doc.add_page_break()
    doc.add_heading("Recommendations and Next Steps", level=2)
    doc.add_paragraph(
        "Based on the comprehensive evaluation results, the following "
        "strategic recommendations are proposed for the upcoming fiscal year:"
    )
    doc.add_paragraph("Expand the Engineering department by 15% to support the growing Aurora platform and related product initiatives.", style="List Number")
    doc.add_paragraph("Allocate additional budget for Marketing's digital transformation, targeting a further 20% improvement in conversion rates.", style="List Number")
    doc.add_paragraph("Implement the Sales Excellence program company-wide, extending the consultative methodology to partner channels.", style="List Number")
    doc.add_paragraph("Continue investment in R&D innovation lab with focus on AI-driven product enhancements and automation.", style="List Number")
    doc.add_paragraph("Enhance the employee wellness program with on-site health screenings and expanded mental health support.", style="List Number")
    doc.add_paragraph("")
    doc.add_paragraph(
        "This report will be reviewed at the quarterly board meeting "
        "scheduled for April 22, 2025. Department heads are requested to "
        "submit detailed implementation plans for their respective areas "
        "by April 8, 2025."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Respectfully submitted,")
    doc.add_paragraph("Dr. Amanda Richardson")
    doc.add_paragraph("Chief Human Resources Officer")
    doc.add_paragraph("Meridian Technologies Inc.")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
