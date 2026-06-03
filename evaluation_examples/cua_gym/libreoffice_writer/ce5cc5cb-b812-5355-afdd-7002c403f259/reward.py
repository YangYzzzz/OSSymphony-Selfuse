"""
Reward Script: Insert page number field in footer with uppercase Roman numerals
Task ID: writer_fs_074
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Page number format set to uppercase Roman (pgNumType fmt="upperRoman")
  Component 2 (0.35): Footer contains a PAGE field code (w:instrText with PAGE)
  Component 3 (0.30): Footer paragraph is centered
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_074'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # We only need to check section 0 (single-section document)
    if len(doc.sections) == 0:
        print("FAIL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: Page numbering format is uppercase Roman (0.35 points)
    # The task requires pgNumType with fmt="upperRoman" in the section properties.
    # This element does NOT exist in the initial env, so it differentiates initial vs golden.
    try:
        sect_pr = section._sectPr
        pg_num_type = sect_pr.find(qn('w:pgNumType'))
        if pg_num_type is not None:
            fmt = pg_num_type.get(qn('w:fmt'))
            if fmt == 'upperRoman':
                print(f"PASS: Component 1 — pgNumType fmt is 'upperRoman' (0.35 pts)")
                total_score += 0.35
            else:
                print(f"FAIL: Component 1 — pgNumType fmt is '{fmt}', expected 'upperRoman'")
        else:
            print("FAIL: Component 1 — No pgNumType element found in section properties")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Footer contains a PAGE field code (0.35 points)
    # The initial env footer is empty with no field codes.
    # The golden env footer must have w:instrText containing ' PAGE '.
    try:
        footer = section.footer
        has_page_field = False
        for para in footer.paragraphs:
            instr_elements = para._element.findall('.//' + qn('w:instrText'))
            for instr in instr_elements:
                if instr.text and 'PAGE' in instr.text.upper():
                    has_page_field = True
                    break
            if has_page_field:
                break

        if has_page_field:
            print(f"PASS: Component 2 — Footer contains PAGE field code (0.35 pts)")
            total_score += 0.35
        else:
            print("FAIL: Component 2 — No PAGE field code found in footer")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Footer paragraph is centered (0.30 points)
    # The initial env footer paragraph has no alignment set (None).
    # The golden env footer paragraph should be centered.
    try:
        footer = section.footer
        footer_centered = False
        for para in footer.paragraphs:
            # Check if this paragraph has the PAGE field (the one with the page number)
            has_field = len(para._element.findall('.//' + qn('w:instrText'))) > 0
            if has_field:
                alignment = para.paragraph_format.alignment
                # WD_PARAGRAPH_ALIGNMENT.CENTER == 1
                if alignment is not None and alignment == 1:
                    footer_centered = True
                    break

        if footer_centered:
            print(f"PASS: Component 3 — Footer page number paragraph is centered (0.30 pts)")
            total_score += 0.30
        else:
            # Also check if any footer paragraph is centered (in case field detection varies)
            any_centered = False
            for para in footer.paragraphs:
                if para.paragraph_format.alignment is not None and para.paragraph_format.alignment == 1:
                    if para.text.strip():  # has some content
                        any_centered = True
                        break
            if any_centered:
                print(f"PASS: Component 3 — Footer paragraph with content is centered (0.30 pts)")
                total_score += 0.30
            else:
                print("FAIL: Component 3 — Footer page number paragraph is not centered")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice state before verification
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
