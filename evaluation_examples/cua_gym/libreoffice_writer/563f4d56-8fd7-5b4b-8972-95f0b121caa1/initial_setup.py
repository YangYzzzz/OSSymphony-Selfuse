"""
Initial Setup: Create a Writer document with realistic API documentation content
but default/empty document properties (metadata).
Task ID: writer_tech_081
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_081'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document content: CloudSync API Reference (but NO metadata set) ---

    # Title heading
    heading = doc.add_heading('CloudSync API Reference', level=0)

    # Introduction
    intro = doc.add_paragraph()
    intro.add_run('Overview').bold = True
    doc.add_paragraph(
        'The CloudSync REST API provides programmatic access to CloudSync resources. '
        'This document covers authentication, endpoint specifications, request/response '
        'formats, and error handling for version 2.1 of the API.'
    )

    # Section: Authentication
    doc.add_heading('Authentication', level=1)
    doc.add_paragraph(
        'All API requests require a valid Bearer token in the Authorization header. '
        'Tokens can be obtained through the OAuth 2.0 authorization code flow or '
        'client credentials grant.'
    )

    auth_para = doc.add_paragraph()
    run = auth_para.add_run('Authorization: Bearer <your_access_token>')
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_paragraph(
        'Tokens expire after 3600 seconds. Use the refresh token endpoint to obtain '
        'a new access token without re-authentication.'
    )

    # Section: Base URL
    doc.add_heading('Base URL', level=1)
    base_url = doc.add_paragraph()
    run = base_url.add_run('https://api.cloudsync.io/v2.1')
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # Section: Endpoints
    doc.add_heading('Endpoints', level=1)

    # Files endpoint
    doc.add_heading('Files', level=2)

    # Endpoint table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Method', 'Endpoint', 'Description']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    endpoints = [
        ['GET', '/files', 'List all files in the workspace'],
        ['GET', '/files/{id}', 'Retrieve file metadata by ID'],
        ['POST', '/files/upload', 'Upload a new file'],
        ['DELETE', '/files/{id}', 'Delete a file permanently'],
    ]
    for r, row_data in enumerate(endpoints, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')  # spacer

    # Sync endpoint
    doc.add_heading('Sync Operations', level=2)

    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table2.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    sync_endpoints = [
        ['POST', '/sync/start', 'Initiate a sync operation'],
        ['GET', '/sync/status/{job_id}', 'Check sync job status'],
        ['POST', '/sync/cancel/{job_id}', 'Cancel a running sync'],
    ]
    for r, row_data in enumerate(sync_endpoints, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    # Section: Error Handling
    doc.add_heading('Error Handling', level=1)
    doc.add_paragraph(
        'The API returns standard HTTP status codes. Error responses include a JSON '
        'body with an error code and human-readable message.'
    )

    error_table = doc.add_table(rows=5, cols=2)
    error_table.style = 'Table Grid'
    error_headers = ['Status Code', 'Meaning']
    for i, h in enumerate(error_headers):
        cell = error_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    errors = [
        ['400', 'Bad Request - Invalid parameters'],
        ['401', 'Unauthorized - Missing or invalid token'],
        ['404', 'Not Found - Resource does not exist'],
        ['429', 'Too Many Requests - Rate limit exceeded'],
    ]
    for r, row_data in enumerate(errors, 1):
        for c, val in enumerate(row_data):
            error_table.cell(r, c).text = val

    # Section: Rate Limits
    doc.add_heading('Rate Limits', level=1)
    doc.add_paragraph(
        'API requests are rate-limited per account. The default limits are:'
    )
    doc.add_paragraph('1000 requests per minute for read operations', style='List Bullet')
    doc.add_paragraph('100 requests per minute for write operations', style='List Bullet')
    doc.add_paragraph('10 concurrent sync jobs per workspace', style='List Bullet')

    doc.add_paragraph(
        'When a rate limit is exceeded, the API returns HTTP 429 with a '
        'Retry-After header indicating the number of seconds to wait.'
    )

    # Section: Changelog
    doc.add_heading('Changelog', level=1)
    doc.add_paragraph('Version 2.1 (March 2025)').runs[0].bold = True
    doc.add_paragraph('Added sync cancellation endpoint', style='List Bullet')
    doc.add_paragraph('Improved file upload performance for large files', style='List Bullet')
    doc.add_paragraph('Added pagination support for file listing', style='List Bullet')

    doc.add_paragraph('Version 2.0 (January 2025)').runs[0].bold = True
    doc.add_paragraph('Redesigned authentication flow with OAuth 2.0', style='List Bullet')
    doc.add_paragraph('Added real-time sync status tracking', style='List Bullet')
    doc.add_paragraph('Deprecated legacy API keys', style='List Bullet')

    # Clear default author set by python-docx library
    doc.core_properties.author = ''

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
