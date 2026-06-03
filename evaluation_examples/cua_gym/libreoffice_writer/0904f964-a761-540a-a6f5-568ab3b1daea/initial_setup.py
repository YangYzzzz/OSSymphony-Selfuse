"""
Initial Setup: Create a formal business letter document (no envelope page)
Task ID: writer_pd_039
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_039'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup: standard letter size ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Letterhead / Sender Info ---
    sender = doc.add_paragraph()
    sender.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    lines = [
        "Pinnacle Solutions Inc.",
        "100 Innovation Drive, Suite 500",
        "Austin, TX 78701",
        "Phone: (512) 555-0198",
        "Email: contact@pinnaclesolutions.com",
    ]
    for i, line in enumerate(lines):
        run = sender.add_run(line)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        if i < len(lines) - 1:
            sender.add_run("\n").font.size = Pt(10)

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(24)
    run = date_para.add_run("March 28, 2026")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # --- Recipient Address ---
    recipient = doc.add_paragraph()
    recipient.paragraph_format.space_before = Pt(12)
    recipient_lines = [
        "Mr. James Morrison",
        "VP Operations",
        "Horizon Enterprises",
        "2500 Corporate Blvd",
        "Denver, CO 80202",
    ]
    for i, line in enumerate(recipient_lines):
        run = recipient.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if i < len(recipient_lines) - 1:
            recipient.add_run("\n").font.size = Pt(11)

    # --- Salutation ---
    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_before = Pt(12)
    run = salutation.add_run("Dear Mr. Morrison,")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # --- Body Paragraphs (enough for 2 pages) ---
    body_texts = [
        "Thank you for taking the time to meet with our team last Thursday to discuss the "
        "proposed partnership between Pinnacle Solutions and Horizon Enterprises. We were "
        "impressed by your organization's commitment to operational excellence and believe "
        "there is a strong alignment between our strategic objectives.",

        "As discussed during our meeting, Pinnacle Solutions is prepared to offer a comprehensive "
        "suite of enterprise resource planning tools tailored specifically to the logistics and "
        "supply chain management challenges that Horizon Enterprises currently faces. Our platform "
        "has been successfully deployed across 47 Fortune 500 companies and consistently delivers "
        "a 23% improvement in operational efficiency within the first fiscal quarter of implementation.",

        "The proposed engagement would consist of three phases. Phase One involves a detailed "
        "assessment of your current infrastructure, including interviews with department heads "
        "across your Denver, Seattle, and Chicago offices. This typically takes four to six weeks "
        "and results in a comprehensive gap analysis report. Phase Two focuses on system design "
        "and customization, where our engineering team works closely with your IT department to "
        "build the integration layer between our platform and your existing SAP environment.",

        "Phase Three covers deployment, training, and ongoing support. We assign a dedicated "
        "account manager and a team of three implementation specialists who remain on-site for "
        "the first 90 days post-launch. Our standard service level agreement guarantees 99.7% "
        "uptime and a four-hour response window for any critical issues.",

        "Regarding the financial terms, I have enclosed a preliminary cost estimate based on "
        "the scope we discussed. The total investment for the first year, including licensing, "
        "customization, and support, would be approximately $2.4 million. This represents a "
        "15% discount from our standard enterprise pricing, reflecting our interest in building "
        "a long-term relationship with Horizon Enterprises.",

        "I would also like to highlight several case studies from clients in similar industries. "
        "TransGlobal Logistics saw a 31% reduction in order processing time after implementing "
        "our platform. Pacific Rim Distributors reported annual savings of $1.8 million in "
        "inventory carrying costs. NorthStar Manufacturing achieved full ROI within 14 months "
        "of deployment, significantly ahead of the projected 18-month timeline.",

        "We understand that a decision of this magnitude requires careful consideration by your "
        "executive team and board of directors. To facilitate your evaluation process, we would "
        "be happy to arrange a demonstration session at your Denver headquarters. Our VP of "
        "Client Relations, Dr. Angela Wei, would personally conduct the presentation and "
        "address any technical questions your team may have.",

        "Additionally, we can connect you with reference contacts at three of our existing "
        "clients who have agreed to share their experiences. These conversations often prove "
        "invaluable in understanding the real-world impact of our solutions beyond what any "
        "proposal document can convey.",

        "Please do not hesitate to reach out if you have any questions or if there is additional "
        "information we can provide. I am available at your convenience for a follow-up call or "
        "meeting. My direct line is (512) 555-0198, extension 401, and I typically respond to "
        "emails within the same business day.",
    ]

    for text in body_texts:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    # --- Closing ---
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(18)
    run = closing.add_run("Sincerely,")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Signature block
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(36)
    sig_lines = [
        "Robert A. Calloway",
        "Senior Vice President, Business Development",
        "Pinnacle Solutions Inc.",
    ]
    for i, line in enumerate(sig_lines):
        run = sig.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if i < len(sig_lines) - 1:
            sig.add_run("\n").font.size = Pt(11)

    # --- Enclosure note ---
    enc = doc.add_paragraph()
    enc.paragraph_format.space_before = Pt(12)
    run = enc.add_run("Enclosures: Preliminary Cost Estimate, Case Study Summaries (3)")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
