"""
Reward Script: Envelope addressing setup in Business_Letter.docx
Task ID: writer_pd_039
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Envelope section exists with #10 dimensions (9.5 x 4.125 inches)
  Component 2 (0.25): Return address in top-left with correct content and ~10pt font
  Component 3 (0.25): Delivery address centered with correct content and ~12pt font
  Component 4 (0.20): Original letter content preserved in subsequent section
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Emu

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_039'

# Tolerance for dimension checks (0.25 inch tolerance)
DIM_TOLERANCE = Inches(0.25)

# Tolerance for font size checks (1pt tolerance)
FONT_TOLERANCE = Pt(1)


def verify_task(file_path):
    """
    Verify envelope addressing setup with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 2 sections (envelope + letter)
    num_sections = len(doc.sections)
    if num_sections < 2:
        print(f"FAIL: Document has {num_sections} section(s), need at least 2 (envelope + letter)")
        print("REWARD: 0.0")
        return 0.0

    envelope_section = doc.sections[0]
    letter_section = doc.sections[1]

    # Component 1: Envelope section has #10 dimensions (9.5 x 4.125 inches) (0.30 points)
    try:
        expected_width = Inches(9.5)
        expected_height = Inches(4.125)
        actual_width = envelope_section.page_width
        actual_height = envelope_section.page_height

        width_ok = abs(actual_width - expected_width) <= DIM_TOLERANCE
        height_ok = abs(actual_height - expected_height) <= DIM_TOLERANCE

        if width_ok and height_ok:
            print(f"PASS: Component 1 — Envelope dimensions correct: "
                  f"{actual_width/914400:.3f}x{actual_height/914400:.3f} in (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 1 — Envelope dimensions wrong: "
                  f"{actual_width/914400:.3f}x{actual_height/914400:.3f} in, "
                  f"expected ~9.500x4.125 in")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Return address in top-left area with correct content and ~10pt font (0.25 points)
    try:
        # Find the first non-empty paragraph in the envelope section
        # Paragraphs belonging to section 0 are before the section break
        # In the golden doc, Para 0 is the return address
        envelope_paras = []
        for para in doc.paragraphs:
            # Check if this paragraph belongs to envelope section
            # Paragraphs before the first section break belong to section 0
            envelope_paras.append(para)
            # Check if this paragraph contains a section break (end of section)
            sect_pr = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
            ppr_sect = para._element.findall(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr/'
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
            if sect_pr or ppr_sect:
                break

        # Also check body-level sectPr approach: paragraphs up to section break
        # Simpler: find paragraphs with return address content
        return_addr_found = False
        return_font_ok = False

        # Expected return address components
        return_addr_parts = ["Pinnacle Solutions Inc.", "100 Innovation Drive", "Suite 500", "Austin", "TX 78701"]

        for para in envelope_paras:
            text = para.text.strip()
            if not text:
                continue
            # Check if this paragraph contains return address
            if "Pinnacle Solutions" in text and "Austin" in text:
                return_addr_found = True
                # Check all required parts
                all_parts_found = all(part in text for part in return_addr_parts)
                if not all_parts_found:
                    print(f"FAIL: Component 2 — Return address missing parts. Found: {text[:100]}")
                    break

                # Check font size (~10pt = 127000 EMU, tolerance 1pt)
                font_sizes = []
                for run in para.runs:
                    if run.font.size is not None and run.text.strip():
                        font_sizes.append(run.font.size)

                if font_sizes:
                    avg_size = sum(font_sizes) / len(font_sizes)
                    target_size = Pt(10)
                    if abs(avg_size - target_size) <= FONT_TOLERANCE:
                        return_font_ok = True
                    else:
                        print(f"FAIL: Component 2 — Return address font size "
                              f"{avg_size/12700:.1f}pt, expected ~10pt")
                else:
                    # If no explicit font size, it might be inherited - still count content
                    return_font_ok = False

                # Check alignment is LEFT (or default)
                align = para.paragraph_format.alignment
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                if align is None or align == WD_PARAGRAPH_ALIGNMENT.LEFT:
                    pass  # Good - left aligned
                else:
                    print(f"FAIL: Component 2 — Return address alignment is {align}, expected LEFT")
                    return_addr_found = False
                break

        if return_addr_found and all_parts_found:
            sub_score = 0.0
            sub_score += 0.15  # content correct
            if return_font_ok:
                sub_score += 0.10  # font size correct
            else:
                print(f"PARTIAL: Component 2 — Return address content OK but font size off")
            print(f"PASS: Component 2 — Return address verified ({sub_score:.2f} pts)")
            total_score += sub_score
        elif return_addr_found:
            print(f"FAIL: Component 2 — Return address found but incomplete")
        else:
            print(f"FAIL: Component 2 — Return address not found in envelope section")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Delivery address centered with correct content and ~12pt font (0.25 points)
    try:
        delivery_addr_found = False
        delivery_centered = False
        delivery_font_ok = False

        delivery_addr_parts = ["Mr. James Morrison", "VP Operations", "Horizon Enterprises",
                               "2500 Corporate Blvd", "Denver", "CO 80202"]

        for para in envelope_paras:
            text = para.text.strip()
            if not text:
                continue
            # Check if this paragraph contains delivery address
            if "James Morrison" in text and "Denver" in text:
                delivery_addr_found = True
                all_delivery_parts = all(part in text for part in delivery_addr_parts)
                if not all_delivery_parts:
                    print(f"FAIL: Component 3 — Delivery address missing parts. Found: {text[:100]}")
                    break

                # Check centering
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                align = para.paragraph_format.alignment
                if align == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    delivery_centered = True
                else:
                    print(f"FAIL: Component 3 — Delivery address not centered (align={align})")

                # Check font size (~12pt = 152400 EMU, tolerance 1pt)
                font_sizes = []
                for run in para.runs:
                    if run.font.size is not None and run.text.strip():
                        font_sizes.append(run.font.size)

                if font_sizes:
                    avg_size = sum(font_sizes) / len(font_sizes)
                    target_size = Pt(12)
                    if abs(avg_size - target_size) <= FONT_TOLERANCE:
                        delivery_font_ok = True
                    else:
                        print(f"FAIL: Component 3 — Delivery address font size "
                              f"{avg_size/12700:.1f}pt, expected ~12pt")
                break

        if delivery_addr_found and all_delivery_parts:
            sub_score = 0.0
            sub_score += 0.10  # content correct
            if delivery_centered:
                sub_score += 0.10  # centered
            if delivery_font_ok:
                sub_score += 0.05  # font size correct
            print(f"PASS: Component 3 — Delivery address verified ({sub_score:.2f} pts)")
            total_score += sub_score
        elif delivery_addr_found:
            print(f"FAIL: Component 3 — Delivery address found but incomplete")
        else:
            print(f"FAIL: Component 3 — Delivery address not found in envelope section")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Original letter content preserved in section 1+ (0.20 points)
    try:
        # The letter section should contain the original business letter content
        # Key markers: "Dear Mr. Morrison", "Sincerely", "Robert A. Calloway"
        # These should appear AFTER the envelope section

        # Get all paragraph texts after envelope section
        in_letter = False
        letter_texts = []
        for para in doc.paragraphs:
            if in_letter:
                letter_texts.append(para.text.strip())
            # Detect section boundary: paragraphs after envelope belong to letter
            sect_pr = para._element.findall(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr/'
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
            if sect_pr and not in_letter:
                in_letter = True

        # If the section break is at body level, use different approach
        if not letter_texts:
            # Fallback: look for letter markers in paragraphs after the delivery address
            found_delivery = False
            for para in doc.paragraphs:
                if found_delivery:
                    letter_texts.append(para.text.strip())
                if "James Morrison" in para.text and "Denver" in para.text:
                    found_delivery = True

        full_letter_text = ' '.join(letter_texts)

        markers = ["Dear Mr. Morrison", "Sincerely", "Robert A. Calloway",
                    "Pinnacle Solutions", "proposed engagement"]
        markers_found = sum(1 for m in markers if m in full_letter_text)

        if markers_found >= 4:
            print(f"PASS: Component 4 — Original letter content preserved "
                  f"({markers_found}/{len(markers)} markers found) (0.20 pts)")
            total_score += 0.20
        elif markers_found >= 2:
            partial = round(0.20 * markers_found / len(markers), 2)
            print(f"PARTIAL: Component 4 — Letter partially preserved "
                  f"({markers_found}/{len(markers)} markers) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Original letter content missing "
                  f"({markers_found}/{len(markers)} markers found)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice edits before verification
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
