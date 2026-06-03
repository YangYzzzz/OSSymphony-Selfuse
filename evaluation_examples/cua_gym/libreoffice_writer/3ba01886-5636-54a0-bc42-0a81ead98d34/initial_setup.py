"""
Initial Setup: Create a 12-page Analysis Report with Table 3 on page 2 and
incomplete cross-reference text on page 8.
Task ID: writer_af_030
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_af_030'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add a manual page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_body_text(doc, text, bold_first=False):
    p = doc.add_paragraph()
    if bold_first:
        parts = text.split('.', 1)
        run = p.add_run(parts[0] + '.')
        run.bold = True
        if len(parts) > 1:
            p.add_run(parts[1])
    else:
        p.add_run(text)
    return p


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ===== PAGE 1: Title Page =====
    doc.add_paragraph()  # spacing
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('Meridian Global Consulting', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_heading('Annual Performance Analysis Report', level=1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = info.add_run('Fiscal Year 2024-2025\nPrepared by the Strategic Analytics Division\nConfidential - Internal Use Only')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_p.add_run('Date: March 15, 2025')
    run.font.size = Pt(11)

    add_page_break(doc)

    # ===== PAGE 2: Executive Summary with Table 3 =====
    add_heading_styled(doc, '1. Executive Summary', level=1)

    add_body_text(doc, 'This report provides a comprehensive analysis of Meridian Global Consulting\'s financial performance across all business units for fiscal year 2024-2025. Our organization achieved significant milestones despite a challenging macroeconomic environment characterized by rising interest rates and shifting market dynamics.')

    add_body_text(doc, 'Key performance indicators demonstrate sustained growth in our core consulting practice, with notable expansion in the technology advisory segment. Revenue from digital transformation engagements increased by 23% year-over-year, validating our strategic pivot towards technology-enabled consulting solutions.')

    # Table 1 caption
    cap1 = doc.add_paragraph()
    run = cap1.add_run('Table 1: Revenue Overview by Region')
    run.bold = True
    run.font.size = Pt(10)

    table1 = doc.add_table(rows=5, cols=4)
    table1.style = 'Table Grid'
    headers1 = ['Region', 'Q1 ($M)', 'Q2 ($M)', 'Annual ($M)']
    for i, h in enumerate(headers1):
        table1.cell(0, i).text = h
    data1 = [
        ['North America', '45.2', '48.7', '187.4'],
        ['Europe', '32.1', '34.5', '132.8'],
        ['Asia Pacific', '18.9', '21.3', '79.6'],
        ['Latin America', '8.4', '9.1', '35.2'],
    ]
    for r, row_data in enumerate(data1, 1):
        for c, val in enumerate(row_data):
            table1.cell(r, c).text = val

    add_body_text(doc, 'The North American market continues to be our largest revenue contributor, accounting for approximately 43% of total global revenue. European operations showed resilient growth despite regulatory headwinds in several key markets.')

    # Table 2 caption
    cap2 = doc.add_paragraph()
    run = cap2.add_run('Table 2: Client Satisfaction Metrics')
    run.bold = True
    run.font.size = Pt(10)

    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Table Grid'
    headers2 = ['Metric', 'FY2024', 'FY2025']
    for i, h in enumerate(headers2):
        table2.cell(0, i).text = h
    data2 = [
        ['Overall Satisfaction', '4.2/5.0', '4.5/5.0'],
        ['Net Promoter Score', '62', '71'],
        ['Repeat Engagement Rate', '78%', '84%'],
    ]
    for r, row_data in enumerate(data2, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    add_body_text(doc, 'Client satisfaction metrics improved notably, with our Net Promoter Score rising from 62 to 71. This improvement is attributed to the implementation of our enhanced client engagement framework and the introduction of dedicated account management teams for our top-tier clients.')

    # Table 3 caption (THE KEY TABLE for cross-reference)
    cap3 = doc.add_paragraph()
    run = cap3.add_run('Table 3: Quarterly Results')
    run.bold = True
    run.font.size = Pt(10)

    table3 = doc.add_table(rows=6, cols=5)
    table3.style = 'Table Grid'
    headers3 = ['Business Unit', 'Q1 ($M)', 'Q2 ($M)', 'Q3 ($M)', 'Q4 ($M)']
    for i, h in enumerate(headers3):
        table3.cell(0, i).text = h
    data3 = [
        ['Strategy Consulting', '12.4', '13.1', '14.8', '15.2'],
        ['Technology Advisory', '18.7', '20.3', '22.1', '24.6'],
        ['Financial Services', '9.3', '10.1', '10.8', '11.4'],
        ['Healthcare Practice', '7.8', '8.2', '8.9', '9.5'],
        ['Total', '48.2', '51.7', '56.6', '60.7'],
    ]
    for r, row_data in enumerate(data3, 1):
        for c, val in enumerate(row_data):
            table3.cell(r, c).text = val

    add_page_break(doc)

    # ===== PAGE 3: Market Analysis =====
    add_heading_styled(doc, '2. Market Analysis', level=1)

    add_body_text(doc, 'The global consulting market experienced a period of transformation during fiscal year 2024-2025. Industry-wide consolidation accelerated, with three major mergers reshaping the competitive landscape. Meanwhile, boutique firms continued to gain market share in specialized domains such as ESG consulting and artificial intelligence strategy.')

    add_body_text(doc, 'Our competitive positioning analysis reveals that Meridian Global Consulting maintained its ranking among the top fifteen management consulting firms globally. The firm\'s market share in North America increased from 2.3% to 2.7%, while our European market share remained stable at 1.8%. This performance is particularly noteworthy given the aggressive pricing strategies deployed by several larger competitors.')

    add_heading_styled(doc, '2.1 Industry Trends', level=2)

    add_body_text(doc, 'Several macro trends shaped the consulting industry during the reporting period. The accelerated adoption of generative AI technologies created unprecedented demand for AI strategy consulting. Organizations across sectors sought guidance on responsible AI implementation, workforce transformation, and technology governance frameworks.')

    add_body_text(doc, 'Digital transformation engagements continued to dominate the pipeline, representing 38% of all new business opportunities. Cloud migration projects, while maturing, still accounted for 22% of technology consulting revenues. Cybersecurity advisory services experienced the fastest growth rate at 34% year-over-year.')

    add_heading_styled(doc, '2.2 Competitive Landscape', level=2)

    add_body_text(doc, 'The merger of Henderson & Associates with Pacific Consulting Group created a formidable competitor in the Asia-Pacific region. However, their integration challenges provided an opportunity for Meridian to capture several key accounts in Singapore, Tokyo, and Sydney. Our rapid response team successfully onboarded twelve new enterprise clients during the transition period.')

    add_body_text(doc, 'Pricing pressures intensified in the mid-market segment, where automated advisory platforms began displacing traditional consulting engagements for standardized assessments and benchmarking exercises. In response, Meridian invested heavily in proprietary analytical tools and frameworks to differentiate our service delivery model.')

    add_page_break(doc)

    # ===== PAGE 4: Financial Performance =====
    add_heading_styled(doc, '3. Financial Performance', level=1)

    add_body_text(doc, 'Total revenue for fiscal year 2024-2025 reached $435.0 million, representing a 12.4% increase over the prior year. This growth was driven primarily by the technology advisory practice, which contributed an incremental $18.3 million in revenue compared to the previous fiscal year.')

    add_heading_styled(doc, '3.1 Revenue Analysis', level=2)

    add_body_text(doc, 'Revenue composition shifted meaningfully during the year. Recurring advisory retainer agreements now represent 41% of total revenue, up from 34% in the prior year. This shift toward predictable revenue streams improves our financial visibility and supports more strategic resource allocation decisions. Project-based engagements, while declining as a percentage of revenue, still grew in absolute terms.')

    add_body_text(doc, 'The average engagement size increased by 15%, reflecting our deliberate strategy to pursue larger, more complex transformation programs. The number of engagements exceeding $5 million grew from 23 to 31, with our largest single engagement valued at $18.7 million for a comprehensive digital transformation program with a multinational pharmaceutical company.')

    add_heading_styled(doc, '3.2 Profitability Metrics', level=2)

    add_body_text(doc, 'Gross margin improved to 52.3% from 49.8%, driven by better utilization rates and a favorable mix shift toward higher-margin advisory work. The firm achieved an operating margin of 18.7%, exceeding our target of 17.5%. EBITDA reached $94.2 million, a record for the organization.')

    add_body_text(doc, 'Partner compensation and profit distribution totaled $67.8 million, representing a 14% increase per equity partner. Associate and manager bonus pools were funded at 110% of target, reflecting the strong performance across all practice areas. We also increased our investment in employee development and wellness programs by 22%.')

    add_page_break(doc)

    # ===== PAGE 5: Practice Area Deep Dives =====
    add_heading_styled(doc, '4. Practice Area Performance', level=1)

    add_heading_styled(doc, '4.1 Strategy Consulting', level=2)

    add_body_text(doc, 'The Strategy Consulting practice delivered $55.5 million in revenue, a 9% increase year-over-year. The practice benefited from strong demand for corporate strategy refreshes as organizations adapted to post-pandemic operating models. Notable engagements included a comprehensive market entry strategy for a Fortune 100 technology company expanding into Southeast Asia.')

    add_body_text(doc, 'Our strategy team expanded to 142 professionals, including 18 new hires from top MBA programs and 7 experienced lateral hires from competitor firms. The practice launched two new service lines: Climate Strategy Advisory and Geopolitical Risk Assessment, both of which achieved profitability within their first year of operation.')

    add_heading_styled(doc, '4.2 Technology Advisory', level=2)

    add_body_text(doc, 'Technology Advisory emerged as our fastest-growing practice, generating $85.7 million in revenue, up 27% from the prior year. The practice capitalized on the surge in demand for AI strategy consulting, enterprise architecture modernization, and data analytics implementations. Key client wins included three major financial institutions and two global retailers.')

    add_body_text(doc, 'The practice invested significantly in developing proprietary frameworks for AI readiness assessment and technology due diligence. Our TechPulse platform, launched in Q2, has been adopted by 45 clients for continuous technology landscape monitoring and strategic planning support.')

    add_heading_styled(doc, '4.3 Financial Services', level=2)

    add_body_text(doc, 'The Financial Services practice contributed $41.6 million in revenue, growing 11% year-over-year. Regulatory advisory services were particularly strong, driven by evolving compliance requirements around digital assets, open banking, and cross-border payment regulations. The practice secured a landmark multi-year advisory contract with one of Europe\'s largest banking groups.')

    add_page_break(doc)

    # ===== PAGE 6: Human Capital =====
    add_heading_styled(doc, '5. Human Capital and Organizational Development', level=1)

    add_body_text(doc, 'Meridian Global Consulting employed 1,847 professionals across 23 offices worldwide at fiscal year-end. Headcount grew 8% during the year, with the most significant additions in the technology advisory and healthcare practices. Our employee retention rate remained strong at 87%, well above the industry average of 79%.')

    add_body_text(doc, 'The firm continued its commitment to diversity, equity, and inclusion. Women now represent 44% of our professional staff and 31% of our partnership. We achieved our goal of having at least 25% of new partner promotions go to candidates from underrepresented groups. Our Employee Resource Groups expanded from six to nine, adding groups focused on veterans, neurodiversity, and sustainability champions.')

    add_heading_styled(doc, '5.1 Talent Acquisition', level=2)

    add_body_text(doc, 'Campus recruiting remained competitive, with Meridian receiving over 12,000 applications for 180 entry-level positions. Our acceptance rate among offered candidates improved to 82%, reflecting the firm\'s strengthening employer brand. Notable recruiting achievements included establishing new partnerships with three additional target universities in Europe and Asia.')

    add_body_text(doc, 'Lateral hiring activity increased substantially, with 67 experienced professionals joining from competitor firms, corporate strategy functions, and technology companies. This influx of experienced talent has strengthened our capabilities in emerging areas such as quantum computing strategy, metaverse advisory, and sustainable supply chain optimization.')

    add_heading_styled(doc, '5.2 Professional Development', level=2)

    add_body_text(doc, 'Investment in professional development reached $8.4 million, a 22% increase from the prior year. Each professional completed an average of 94 hours of formal training, supplemented by mentoring, coaching, and experiential learning programs. Our Meridian Leadership Academy graduated its fifth cohort of 28 high-potential senior managers.')

    add_body_text(doc, 'The firm launched its Digital Skills Initiative, providing all professionals with foundational training in data analytics, AI/ML concepts, and digital tool proficiency. Completion rates exceeded 90% within the first six months, and feedback surveys indicated high satisfaction with the program content and delivery format.')

    add_page_break(doc)

    # ===== PAGE 7: Client Engagement Highlights =====
    add_heading_styled(doc, '6. Client Engagement Highlights', level=1)

    add_body_text(doc, 'During fiscal year 2024-2025, Meridian Global Consulting delivered exceptional results across a diverse portfolio of client engagements. Several projects exemplified our commitment to creating measurable value for our clients and pushing the boundaries of consulting excellence.')

    add_heading_styled(doc, '6.1 Project Phoenix - Global Retailer Transformation', level=2)

    add_body_text(doc, 'Meridian led a comprehensive digital transformation program for one of North America\'s largest retail chains, encompassing supply chain optimization, customer experience redesign, and workforce enablement. The 18-month engagement involved a team of 35 consultants and resulted in a documented $230 million in annual cost savings and a 15% improvement in customer satisfaction scores.')

    add_heading_styled(doc, '6.2 Project Atlas - Financial Services Integration', level=2)

    add_body_text(doc, 'Following the acquisition of a mid-sized investment bank by a global financial services group, Meridian was engaged to design and execute the post-merger integration strategy. Our team developed a detailed 100-day plan, managed cultural alignment workshops, and oversaw the technology platform consolidation. The integration was completed three months ahead of schedule, with synergy realization exceeding initial projections by 18%.')

    add_heading_styled(doc, '6.3 Project Horizon - Healthcare System Redesign', level=2)

    add_body_text(doc, 'A major healthcare system engaged Meridian to redesign its care delivery model across 14 hospitals and 200 outpatient facilities. Our team developed a patient-centered care framework that reduced average length of stay by 1.2 days, improved readmission rates by 23%, and generated $45 million in operational savings. The project received the Healthcare Advisory Board\'s Innovation Award.')

    add_body_text(doc, 'These engagements demonstrate Meridian\'s ability to deliver transformative results across sectors and geographies, leveraging our deep expertise, proprietary methodologies, and commitment to measurable outcomes.')

    add_page_break(doc)

    # ===== PAGE 8: Strategic Outlook (with incomplete cross-reference) =====
    add_heading_styled(doc, '7. Strategic Outlook and Recommendations', level=1)

    add_body_text(doc, 'Looking ahead to fiscal year 2025-2026, Meridian Global Consulting is well-positioned to capitalize on several emerging opportunities while navigating anticipated market challenges. Our strategic planning committee has identified five priority areas for investment and growth.')

    add_body_text(doc, 'First, we will continue to scale our AI and automation advisory capabilities, targeting a 40% revenue increase in this segment. Second, geographic expansion into the Middle East and Africa will be supported by the opening of new offices in Dubai and Nairobi. Third, we plan to acquire a specialized data engineering firm to strengthen our technology delivery capabilities.')

    add_body_text(doc, 'Fourth, our client engagement model will evolve to incorporate more outcome-based pricing structures, aligning our incentives more closely with client value creation. Fifth, investment in proprietary technology platforms will increase by 30%, with a focus on AI-powered analytics tools that enhance our consultants\' productivity and insight generation capabilities.')

    add_body_text(doc, 'The quarterly financial trajectory provides strong confidence in our growth outlook. As shown in')

    add_body_text(doc, 'Our strategic investments in talent, technology, and market expansion position Meridian for sustained growth and continued leadership in the global consulting market. The management team is confident in achieving our fiscal year 2025-2026 revenue target of $500 million.')

    add_page_break(doc)

    # ===== PAGE 9: Risk Management =====
    add_heading_styled(doc, '8. Risk Management and Compliance', level=1)

    add_body_text(doc, 'Meridian Global Consulting maintains a robust enterprise risk management framework that identifies, assesses, and mitigates risks across all dimensions of our operations. During fiscal year 2024-2025, the Risk Committee met quarterly and conducted two comprehensive risk reviews.')

    add_body_text(doc, 'Key risk categories monitored include operational risk, reputational risk, regulatory compliance risk, cybersecurity risk, and talent retention risk. Our risk appetite statement was updated in Q2 to reflect the evolving threat landscape and the firm\'s increased reliance on digital infrastructure and cloud-based service delivery models.')

    add_heading_styled(doc, '8.1 Cybersecurity Posture', level=2)

    add_body_text(doc, 'The firm invested $3.2 million in cybersecurity enhancements during the year, including implementation of a zero-trust network architecture, deployment of advanced endpoint detection and response tools, and establishment of a 24/7 security operations center. Penetration testing conducted by an independent third party identified no critical vulnerabilities, and all medium-severity findings were remediated within 30 days.')

    add_body_text(doc, 'Employee cybersecurity awareness training was expanded to include monthly phishing simulations and quarterly security workshops. The phishing simulation click-through rate decreased from 12% to 3.5% over the course of the year, demonstrating improved security awareness across the organization.')

    add_heading_styled(doc, '8.2 Regulatory Compliance', level=2)

    add_body_text(doc, 'The firm maintained full compliance with all applicable regulatory requirements across our operating jurisdictions. Notable compliance achievements include successful completion of SOC 2 Type II certification, GDPR compliance audit, and ISO 27001 recertification. No regulatory penalties or sanctions were imposed during the reporting period.')

    add_page_break(doc)

    # ===== PAGE 10: Innovation and Research =====
    add_heading_styled(doc, '9. Innovation and Research', level=1)

    add_body_text(doc, 'The Meridian Innovation Lab, established in 2023, made significant strides in developing next-generation consulting tools and methodologies. The lab\'s annual budget of $5.8 million supported twelve research initiatives across three focus areas: artificial intelligence applications, sustainability analytics, and organizational network analysis.')

    add_body_text(doc, 'Three innovations were commercialized during the year and integrated into client engagements. The AI-powered Market Intelligence Dashboard, the Carbon Footprint Optimization Engine, and the Organizational Health Index tool each received positive client feedback and contributed $4.2 million in incremental revenue through premium service offerings.')

    add_heading_styled(doc, '9.1 Research Partnerships', level=2)

    add_body_text(doc, 'Meridian established formal research partnerships with four leading universities: MIT Sloan School of Management, London Business School, INSEAD, and the National University of Singapore. These partnerships support joint research projects, provide access to cutting-edge academic insights, and strengthen our campus recruiting pipeline.')

    add_body_text(doc, 'Published research output included 15 white papers, 8 journal articles, and 3 book chapters authored by Meridian professionals. The firm\'s Thought Leadership Index ranking improved from 12th to 8th among global consulting firms, reflecting our increasing intellectual contribution to the industry.')

    add_page_break(doc)

    # ===== PAGE 11: Sustainability =====
    add_heading_styled(doc, '10. Sustainability and Corporate Responsibility', level=1)

    add_body_text(doc, 'Meridian Global Consulting is committed to operating responsibly and contributing positively to the communities in which we work. Our sustainability strategy encompasses environmental stewardship, social responsibility, and governance excellence.')

    add_body_text(doc, 'The firm achieved carbon neutrality for Scope 1 and Scope 2 emissions during the fiscal year, meeting our 2025 target one year ahead of schedule. Total carbon emissions decreased by 28% compared to the baseline year, driven by renewable energy procurement, office space optimization, and reduced business travel through virtual engagement models.')

    add_heading_styled(doc, '10.1 Community Engagement', level=2)

    add_body_text(doc, 'Professional staff contributed over 15,000 hours of pro bono consulting services to non-profit organizations, social enterprises, and government agencies. Notable pro bono engagements included developing a digital literacy program for underserved communities in collaboration with a major technology company, and creating a strategic plan for a global health organization focused on pandemic preparedness.')

    add_body_text(doc, 'The Meridian Foundation distributed $2.1 million in grants supporting education, environmental conservation, and community development initiatives across 14 countries. The Foundation\'s flagship scholarship program funded 45 students from disadvantaged backgrounds to pursue graduate education in business and technology.')

    add_page_break(doc)

    # ===== PAGE 12: Appendices =====
    add_heading_styled(doc, 'Appendix A: Methodology Notes', level=1)

    add_body_text(doc, 'Financial data presented in this report has been prepared in accordance with International Financial Reporting Standards (IFRS) and reviewed by our independent auditors, Whitmore & Associates LLP. Revenue recognition follows the percentage-of-completion method for fixed-fee engagements and actual hours billed for time-and-materials contracts.')

    add_body_text(doc, 'Market share data is sourced from the Global Consulting Market Analysis published by Industry Research Associates, supplemented by our internal market intelligence system. Client satisfaction metrics are derived from our annual client survey, which achieved a 72% response rate across the active client base.')

    add_heading_styled(doc, 'Appendix B: Office Locations', level=1)

    offices = [
        'New York (Headquarters)', 'San Francisco', 'Chicago', 'Boston', 'Washington D.C.',
        'Toronto', 'London', 'Frankfurt', 'Paris', 'Zurich', 'Singapore', 'Tokyo',
        'Sydney', 'Mumbai', 'Shanghai', 'Sao Paulo', 'Mexico City', 'Dubai',
        'Johannesburg', 'Seoul', 'Hong Kong', 'Stockholm', 'Milan'
    ]
    for office in offices:
        doc.add_paragraph(office, style='List Bullet')

    add_body_text(doc, 'This report was prepared by the Strategic Analytics Division under the direction of the Chief Strategy Officer. For inquiries, please contact the Office of the Managing Partner.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
