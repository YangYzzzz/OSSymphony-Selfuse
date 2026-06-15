"""
Reward Script: Cross-reference to Table 3 with auto-updating page number
Task ID: writer_af_030
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): REF field code referencing Table 3 caption exists in target paragraph
  Component 2 (0.25): PAGEREF field code for page number exists in target paragraph
  Component 3 (0.25): Full text matches pattern "As shown in Table 3 on page <N>"
  Component 4 (0.25): Bookmark target (_Ref_Table3) exists near the Table 3 caption
"""

import os
import re
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_af_030'
TARGET_PARA_INDEX = 71  # Paragraph containing "As shown in"
CAPTION_PARA_INDEX = 16  # Paragraph containing "Table 3: Quarterly Results"

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': W_NS}


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice changes via Ctrl+S."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def get_field_codes(para_element):
    """Extract field instruction texts from a paragraph XML element.

    Returns a list of instrText content strings found in the paragraph.
    """
    instr_texts = []
    for instr in para_element.findall('.//w:instrText', NS):
        if instr.text:
            instr_texts.append(instr.text.strip())
    return instr_texts


def get_field_display_texts(para_element):
    """Extract display text from field codes (text between separate and end fldChar).

    Returns list of display text strings.
    """
    runs = para_element.findall('.//w:r', NS)
    display_texts = []
    in_field_result = False
    current_text = []

    for run in runs:
        fld_char = run.find('w:fldChar', NS)
        if fld_char is not None:
            fld_type = fld_char.get(f'{{{W_NS}}}fldCharType')
            if fld_type == 'separate':
                in_field_result = True
                current_text = []
            elif fld_type == 'end':
                if in_field_result and current_text:
                    display_texts.append(''.join(current_text))
                in_field_result = False
                current_text = []
            elif fld_type == 'begin':
                in_field_result = False
                current_text = []
        elif in_field_result:
            t_elem = run.find('w:t', NS)
            if t_elem is not None and t_elem.text:
                current_text.append(t_elem.text)

    return display_texts


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has enough paragraphs
    if len(doc.paragraphs) <= TARGET_PARA_INDEX:
        print(f"CRITICAL: Document has only {len(doc.paragraphs)} paragraphs, expected at least {TARGET_PARA_INDEX + 1}")
        print("REWARD: 0.0")
        return 0.0

    target_para = doc.paragraphs[TARGET_PARA_INDEX]
    target_xml = target_para._element

    # Get field codes from the target paragraph
    field_codes = get_field_codes(target_xml)
    field_display = get_field_display_texts(target_xml)

    print(f"INFO: Target para text: {repr(target_para.text)}")
    print(f"INFO: Field codes found: {field_codes}")
    print(f"INFO: Field display texts: {field_display}")

    # Component 1: REF field code referencing Table 3 caption (0.25 points)
    # The REF field should reference the Table 3 bookmark
    try:
        ref_field_found = False
        for fc in field_codes:
            if fc.startswith('REF') and 'Table3' in fc.replace('_', '').replace(' ', ''):
                ref_field_found = True
                break
        if ref_field_found:
            print(f"PASS: Component 1 -- REF field code found for Table 3 (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 -- No REF field code found referencing Table 3. Fields: {field_codes}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: PAGEREF field code for page number (0.25 points)
    # The PAGEREF field should auto-update the page number
    try:
        pageref_found = False
        for fc in field_codes:
            if fc.startswith('PAGEREF') and 'Table3' in fc.replace('_', '').replace(' ', ''):
                pageref_found = True
                break
        if pageref_found:
            print(f"PASS: Component 2 -- PAGEREF field code found for Table 3 (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 -- No PAGEREF field code found. Fields: {field_codes}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Full text matches expected pattern (0.25 points)
    # Para text should read "...As shown in Table 3 on page <N>"
    try:
        para_text = target_para.text
        pattern = r'As shown in\s+Table\s*3\s+on page\s+\d+'
        if re.search(pattern, para_text):
            print(f"PASS: Component 3 -- Text matches 'As shown in Table 3 on page N' pattern (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 3 -- Text does not match pattern. Got: {repr(para_text)}")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Bookmark _Ref_Table3 exists in document (0.25 points)
    # The cross-reference target bookmark must exist near the Table 3 caption
    try:
        body = doc.element.body
        bookmarks = body.findall('.//w:bookmarkStart', NS)
        bookmark_found = False
        for bm in bookmarks:
            name = bm.get(f'{{{W_NS}}}name', '')
            # Accept any bookmark name that references Table3
            if 'Table3' in name.replace('_', '').replace(' ', ''):
                bookmark_found = True
                break
        if bookmark_found:
            print(f"PASS: Component 4 -- Bookmark for Table 3 cross-reference target exists (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 4 -- No bookmark referencing Table 3 found")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
