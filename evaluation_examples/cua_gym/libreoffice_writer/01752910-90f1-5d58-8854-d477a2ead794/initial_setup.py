"""
Initial Setup: Accept all tracked insertions, reject all tracked deletions
Task ID: writer_rm_038
Domain: libreoffice_writer

Creates a Product Catalog document with 20 tracked changes:
  - 11 insertions (new product descriptions and features)
  - 9 deletions (removed outdated entries)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import copy
from datetime import datetime

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_038'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

AUTHOR = "Emily Rodriguez"
DATE = "2026-03-15T10:30:00Z"

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_run_element(text, bold=False, italic=False, font_name=None, font_size=None, color=None):
    """Create a w:r element with optional formatting."""
    r = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{text}</w:t></w:r>')
    rpr_parts = []
    if bold:
        rpr_parts.append('<w:b/>')
    if italic:
        rpr_parts.append('<w:i/>')
    if font_name:
        rpr_parts.append(f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
    if font_size:
        sz = int(font_size * 2)  # half-points
        rpr_parts.append(f'<w:sz w:val="{sz}"/>')
    if color:
        rpr_parts.append(f'<w:color w:val="{color}"/>')
    if rpr_parts:
        rpr_xml = f'<w:rPr {nsdecls("w")}>{"".join(rpr_parts)}</w:rPr>'
        rpr = parse_xml(rpr_xml)
        r.insert(0, rpr)
    return r


def make_ins_element(text, author=AUTHOR, date=DATE, bold=False, italic=False,
                     font_name=None, font_size=None, color=None):
    """Create a w:ins element wrapping a run (tracked insertion)."""
    ins = parse_xml(
        f'<w:ins {nsdecls("w")} w:id="{next(id_gen)}" '
        f'w:author="{author}" w:date="{date}"/>'
    )
    r = make_run_element(text, bold=bold, italic=italic,
                         font_name=font_name, font_size=font_size, color=color)
    ins.append(r)
    return ins


def make_del_element(text, author=AUTHOR, date=DATE, bold=False, italic=False,
                     font_name=None, font_size=None, color=None):
    """Create a w:del element wrapping a deleted run (tracked deletion)."""
    del_elem = parse_xml(
        f'<w:del {nsdecls("w")} w:id="{next(id_gen)}" '
        f'w:author="{author}" w:date="{date}"/>'
    )
    # Deleted runs use w:delText instead of w:t
    r_xml = f'<w:r {nsdecls("w")}><w:delText xml:space="preserve">{text}</w:delText></w:r>'
    r = parse_xml(r_xml)
    rpr_parts = []
    if bold:
        rpr_parts.append('<w:b/>')
    if italic:
        rpr_parts.append('<w:i/>')
    if font_name:
        rpr_parts.append(f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
    if font_size:
        sz = int(font_size * 2)
        rpr_parts.append(f'<w:sz w:val="{sz}"/>')
    if color:
        rpr_parts.append(f'<w:color w:val="{color}"/>')
    if rpr_parts:
        rpr_xml = f'<w:rPr {nsdecls("w")}>{"".join(rpr_parts)}</w:rPr>'
        rpr = parse_xml(rpr_xml)
        r.insert(0, rpr)
    del_elem.append(r)
    return del_elem


def _id_generator():
    i = 100
    while True:
        yield i
        i += 1

id_gen = _id_generator()


def add_paragraph_with_revisions(doc, elements, style=None, alignment=None):
    """
    Add a paragraph to the document body with mixed normal runs and revision elements.
    elements: list of tuples:
      ('text', text, kwargs) - normal run
      ('ins', text, kwargs)  - tracked insertion
      ('del', text, kwargs)  - tracked deletion
    """
    p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')

    # Add paragraph properties if needed
    if style or alignment:
        ppr_parts = []
        if style:
            ppr_parts.append(f'<w:pStyle w:val="{style}"/>')
        if alignment:
            align_map = {'center': 'center', 'left': 'left', 'right': 'right', 'justify': 'both'}
            ppr_parts.append(f'<w:jc w:val="{align_map.get(alignment, alignment)}"/>')
        ppr_xml = f'<w:pPr {nsdecls("w")}>{"".join(ppr_parts)}</w:pPr>'
        ppr = parse_xml(ppr_xml)
        p.insert(0, ppr)

    for elem_type, text, kwargs in elements:
        if elem_type == 'text':
            r = make_run_element(text, **kwargs)
            p.append(r)
        elif elem_type == 'ins':
            ins = make_ins_element(text, **kwargs)
            p.append(ins)
        elif elem_type == 'del':
            de = make_del_element(text, **kwargs)
            p.append(de)

    doc.element.body.append(p)
    return p


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # === Title ===
    title = doc.add_heading('Meridian Electronics — Product Catalog 2026', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')  # spacer

    # === Section 1: Smartphones ===
    doc.add_heading('1. Smartphones', level=1)

    # Normal paragraph
    doc.add_paragraph(
        'Our smartphone lineup features cutting-edge technology designed for '
        'professionals and everyday users alike. Each device undergoes rigorous '
        'quality testing before reaching consumers.'
    )

    # Product 1 - with insertion #1 (new feature added)
    doc.add_heading('1.1 Meridian Nova X12', level=2)
    p_body = doc.element.body

    # Normal text + insertion of new feature
    add_paragraph_with_revisions(doc, [
        ('text', 'The Nova X12 features a 6.7-inch AMOLED display with 120Hz refresh rate, ', {}),
        ('text', 'powered by the Snapdragon 8 Gen 3 processor. ', {}),
        ('ins', 'Now includes advanced AI-powered noise cancellation for crystal-clear calls. ', {}),  # INS 1
        ('text', 'Available in Midnight Black, Arctic Silver, and Ocean Blue.', {}),
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $899 | Storage: 256GB / 512GB | Battery: 5000mAh', {}),
    ])

    # Product 2 - with deletion #1 (outdated spec removed)
    doc.add_heading('1.2 Meridian Pulse Pro', level=2)
    add_paragraph_with_revisions(doc, [
        ('text', 'The Pulse Pro is our mid-range champion, offering exceptional camera quality ', {}),
        ('text', 'with a 108MP main sensor and 8K video recording capability. ', {}),
        ('del', 'Compatible with 3G CDMA networks for legacy carrier support. ', {}),  # DEL 1
        ('text', 'Dual SIM support with eSIM technology.', {}),
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $599 | Storage: 128GB / 256GB | Battery: 4500mAh', {}),
    ])

    # === Section 2: Laptops ===
    doc.add_heading('2. Laptops', level=1)

    doc.add_paragraph(
        'Meridian laptops are engineered for performance and portability. '
        'From ultrabooks to workstations, we cover every professional need.'
    )

    # Product 3 - with insertion #2 and deletion #2
    doc.add_heading('2.1 Meridian AeroBook Ultra', level=2)
    add_paragraph_with_revisions(doc, [
        ('text', 'Weighing just 2.1 lbs, the AeroBook Ultra delivers all-day battery life ', {}),
        ('text', 'with up to 18 hours of usage. ', {}),
        ('ins', 'Features the new Intel Core Ultra 9 processor with integrated NPU for on-device AI workloads. ', {}),  # INS 2
        ('del', 'Ships with Intel Core i5-1135G7 processor. ', {}),  # DEL 2
        ('text', '14-inch 2.8K OLED display with 100% DCI-P3 coverage.', {}),
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $1,299 | RAM: 16GB / 32GB | Storage: 512GB / 1TB NVMe SSD', {}),
    ])

    # Product 4 - with insertion #3
    doc.add_heading('2.2 Meridian ProStation 7', level=2)
    add_paragraph_with_revisions(doc, [
        ('text', 'The ProStation 7 is built for demanding creative workflows — video editing, ', {}),
        ('text', '3D rendering, and machine learning development. ', {}),
        ('text', 'Equipped with NVIDIA RTX 4090 mobile GPU and 64GB DDR5 memory. ', {}),
        ('ins', 'Now supports Thunderbolt 5 with up to 80Gbps bandwidth for external displays and storage arrays. ', {}),  # INS 3
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $2,799 | Display: 16-inch Mini-LED | Weight: 4.8 lbs', {}),
    ])

    # === Section 3: Audio ===
    doc.add_heading('3. Audio Equipment', level=1)

    doc.add_paragraph(
        'Experience sound the way artists intended. Our audio products combine '
        'premium materials with advanced acoustic engineering.'
    )

    # Product 5 - with deletion #3 and insertion #4
    doc.add_heading('3.1 Meridian SoundArc Headphones', level=2)
    add_paragraph_with_revisions(doc, [
        ('text', 'Over-ear wireless headphones with active noise cancellation ', {}),
        ('text', 'and spatial audio support. ', {}),
        ('del', 'Uses Bluetooth 4.2 with aptX codec support. ', {}),  # DEL 3
        ('ins', 'Upgraded to Bluetooth 5.3 with LC3plus codec for lossless wireless audio transmission. ', {}),  # INS 4
        ('text', '40-hour battery life with quick charge (10 min = 3 hours).', {}),
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $349 | Driver: 40mm | Frequency Response: 20Hz–40kHz', {}),
    ])

    # Product 6 - with deletion #4
    doc.add_heading('3.2 Meridian BassCore Speaker', level=2)
    add_paragraph_with_revisions(doc, [
        ('text', 'Portable Bluetooth speaker with 360-degree sound projection ', {}),
        ('text', 'and IP67 water resistance rating. ', {}),
        ('del', 'Includes auxiliary 3.5mm input and micro-USB charging port. ', {}),  # DEL 4
        ('text', 'USB-C fast charging with 20-hour playback on a single charge.', {}),
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $179 | Weight: 1.5 lbs | Dimensions: 7.2 x 3.1 x 3.1 inches', {}),
    ])

    # === Section 4: Wearables ===
    doc.add_heading('4. Wearables', level=1)

    doc.add_paragraph(
        'Stay connected and track your health with our range of smart wearables. '
        'Designed for comfort and durability in any condition.'
    )

    # Product 7 - with insertion #5 and insertion #6
    doc.add_heading('4.1 Meridian FitBand Quantum', level=2)
    add_paragraph_with_revisions(doc, [
        ('text', 'Advanced fitness tracker with continuous heart rate monitoring, ', {}),
        ('text', 'blood oxygen measurement, and sleep analysis. ', {}),
        ('ins', 'New ECG monitoring feature provides medical-grade heart rhythm analysis with FDA clearance. ', {}),  # INS 5
        ('ins', 'Includes fall detection and automatic emergency SOS with GPS location sharing. ', {}),  # INS 6
        ('text', 'Water resistant to 50 meters for swim tracking.', {}),
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $249 | Battery: 7 days | Display: 1.4-inch AMOLED', {}),
    ])

    # Product 8 - with deletion #5
    doc.add_heading('4.2 Meridian ChronoSmart Watch', level=2)
    add_paragraph_with_revisions(doc, [
        ('text', 'Premium smartwatch with titanium case and sapphire crystal display. ', {}),
        ('text', 'Supports contactless payments, music streaming, and turn-by-turn navigation. ', {}),
        ('del', 'Requires companion smartphone running Android 8.0 or iOS 12. ', {}),  # DEL 5
        ('text', 'Standalone LTE connectivity for calls and messages without your phone.', {}),
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $499 | Case Size: 44mm / 47mm | Battery: 3 days with LTE', {}),
    ])

    # === Section 5: Smart Home ===
    doc.add_heading('5. Smart Home', level=1)

    doc.add_paragraph(
        'Transform your living space with intelligent devices that learn your preferences '
        'and automate daily routines for maximum comfort and efficiency.'
    )

    # Product 9 - with insertion #7 and deletion #6
    doc.add_heading('5.1 Meridian HomeHub Controller', level=2)
    add_paragraph_with_revisions(doc, [
        ('text', 'Central smart home hub with voice assistant integration ', {}),
        ('text', '(Alexa, Google Assistant, Siri). ', {}),
        ('ins', 'Matter and Thread protocol support enables seamless interoperability with 500+ smart home brands. ', {}),  # INS 7
        ('del', 'Limited to Zigbee protocol with support for up to 30 connected devices. ', {}),  # DEL 6
        ('text', '7-inch touchscreen display for device management and video calls.', {}),
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $229 | Connectivity: Wi-Fi 6E, Bluetooth 5.3, Thread | Power: AC adapter', {}),
    ])

    # Product 10 - with insertion #8
    doc.add_heading('5.2 Meridian SecureCam Pro', level=2)
    add_paragraph_with_revisions(doc, [
        ('text', '4K outdoor security camera with color night vision and AI-powered ', {}),
        ('text', 'person/vehicle/package detection. Two-way audio with noise reduction. ', {}),
        ('ins', 'Features on-device AI processing eliminating cloud dependency for privacy-sensitive users — all footage stays local. ', {}),  # INS 8
        ('text', 'Weatherproof design rated for -22°F to 122°F operating temperatures.', {}),
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $199 | Resolution: 4K HDR | Storage: microSD up to 256GB or cloud', {}),
    ])

    # === Section 6: Accessories ===
    doc.add_heading('6. Accessories', level=1)

    doc.add_paragraph(
        'Complete your Meridian ecosystem with premium accessories designed '
        'for seamless integration and enhanced productivity.'
    )

    # Product 11 - with deletion #7 and insertion #9
    doc.add_heading('6.1 Meridian PowerDock Station', level=2)
    add_paragraph_with_revisions(doc, [
        ('text', 'Universal docking station with dual monitor support (up to 4K@60Hz). ', {}),
        ('del', 'Provides USB 3.0 Type-A ports and HDMI 1.4 output. ', {}),  # DEL 7
        ('ins', 'Equipped with USB4 ports, dual HDMI 2.1 outputs, and 2.5GbE Ethernet for high-speed connectivity. ', {}),  # INS 9
        ('text', '100W Power Delivery charging for compatible laptops.', {}),
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $189 | Ports: 12 total | Weight: 0.8 lbs', {}),
    ])

    # Product 12 - with deletion #8 and insertion #10
    doc.add_heading('6.2 Meridian ErgoType Keyboard', level=2)
    add_paragraph_with_revisions(doc, [
        ('text', 'Ergonomic split mechanical keyboard with hot-swappable switches ', {}),
        ('text', 'and programmable RGB backlighting. ', {}),
        ('del', 'Connects via USB-A cable only. ', {}),  # DEL 8
        ('ins', 'Tri-mode connectivity: USB-C wired, Bluetooth 5.3, and 2.4GHz wireless dongle with 1ms response time. ', {}),  # INS 10
        ('text', 'Adjustable tenting angles (0°, 7°, 15°) for wrist comfort.', {}),
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $159 | Layout: 75% | Switch: Gateron Pro 3.0', {}),
    ])

    # Product 13 - with deletion #9 and insertion #11
    doc.add_heading('6.3 Meridian UltraCharge Powerbank', level=2)
    add_paragraph_with_revisions(doc, [
        ('text', 'High-capacity portable charger with 25,000mAh battery. ', {}),
        ('del', 'Maximum output 18W with Qualcomm Quick Charge 3.0. ', {}),  # DEL 9
        ('ins', 'Supports 140W USB-C PD 3.1 output capable of charging laptops at full speed. ', {}),  # INS 11
        ('text', 'Charges three devices simultaneously with intelligent power distribution.', {}),
    ])

    add_paragraph_with_revisions(doc, [
        ('text', 'Price: $89 | Capacity: 25,000mAh | Weight: 1.1 lbs', {}),
    ])

    # Footer paragraph
    doc.add_paragraph('')
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_para.add_run('© 2026 Meridian Electronics Inc. All rights reserved.')
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Count revisions for verification
    body = doc.element.body
    ins_count = len(body.findall('.//' + qn('w:ins')))
    del_count = len(body.findall('.//' + qn('w:del')))
    print(f'Tracked changes: {ins_count} insertions, {del_count} deletions')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
