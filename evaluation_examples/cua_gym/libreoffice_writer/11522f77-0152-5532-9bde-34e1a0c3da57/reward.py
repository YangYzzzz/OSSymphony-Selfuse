"""
Reward Script: Formal Quarterly Business Report in LibreOffice Writer
Task ID: writer_wf_001
Domain: libreoffice_writer
Scoring:
  Component 1 — Title page content (0.15)
  Component 2 — Heading 1 main sections (0.25)
  Component 3 — Heading 2 subsections (0.10)
  Component 4 — Financial data table (0.25)
  Component 5 — Table of Contents section (0.10)
  Component 6 — Page numbers in footer (0.15)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_001'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect paragraph info
    all_paras = [(p.style.name if p.style else 'Normal', p.text) for p in doc.paragraphs]
    all_texts = [t for _, t in all_paras]
    heading1_texts = [t for s, t in all_paras if s == 'Heading 1']
    heading2_texts = [t for s, t in all_paras if s == 'Heading 2']

    # ---------- Component 1: Title page content (0.15 points) ----------
    # The golden file has "Nexora Technologies" and "Q3 2025 Performance Report" on the title page.
    # These are task-introduced (blank doc initially).
    try:
        full_text_lower = ' '.join(all_texts).lower()
        has_company = 'nexora technologies' in full_text_lower
        has_report_title = 'q3 2025' in full_text_lower and 'report' in full_text_lower

        if has_company and has_report_title:
            print(f"PASS: Component 1 — Title page has 'Nexora Technologies' and Q3 2025 report title (0.15 pts)")
            total_score += 0.15
        elif has_company or has_report_title:
            print(f"PARTIAL: Component 1 — Only one of company name or report title found (0.07 pts)")
            total_score += 0.07
        else:
            print(f"FAIL: Component 1 — Missing title page content (Nexora Technologies / Q3 2025 report)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------- Component 2: Heading 1 main sections (0.25 points) ----------
    # Task requires sections: Executive Summary, Financial Highlights, Market Analysis, Recommendations
    # All must use Heading 1 style.
    try:
        required_sections = ['executive summary', 'financial highlights', 'market analysis', 'recommendations']
        found_sections = []
        for req in required_sections:
            for h1 in heading1_texts:
                if req in h1.lower():
                    found_sections.append(req)
                    break

        section_count = len(found_sections)
        if section_count == 4:
            print(f"PASS: Component 2 — All 4 required Heading 1 sections found: {found_sections} (0.25 pts)")
            total_score += 0.25
        elif section_count >= 2:
            partial = round(0.25 * section_count / 4, 2)
            print(f"PARTIAL: Component 2 — {section_count}/4 Heading 1 sections found: {found_sections} ({partial} pts)")
            total_score += partial
        elif section_count == 1:
            print(f"PARTIAL: Component 2 — Only 1/4 Heading 1 sections found: {found_sections} (0.06 pts)")
            total_score += 0.06
        else:
            print(f"FAIL: Component 2 — No required Heading 1 sections found. H1 texts: {heading1_texts}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---------- Component 3: Heading 2 subsections (0.10 points) ----------
    # Task says "Use Heading 2 for subsections". Golden has multiple H2s.
    try:
        if len(heading2_texts) >= 2:
            print(f"PASS: Component 3 — {len(heading2_texts)} Heading 2 subsections found (0.10 pts)")
            total_score += 0.10
        elif len(heading2_texts) == 1:
            print(f"PARTIAL: Component 3 — Only 1 Heading 2 subsection found (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 3 — No Heading 2 subsections found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ---------- Component 4: Financial data table (0.25 points) ----------
    # Task requires a table with Revenue, Expenses, Net Profit for July/Aug/Sept.
    # Golden: 4 rows x 4 cols (Metric, July, August, September)
    try:
        tables = doc.tables
        if len(tables) == 0:
            print(f"FAIL: Component 4 — No tables found in document")
        else:
            best_table_score = 0.0
            for ti, table in enumerate(tables):
                table_score = 0.0
                # Collect all cell text
                all_cell_text = ''
                row_count = len(table.rows)
                col_count = len(table.columns)
                for row in table.rows:
                    for cell in row.cells:
                        all_cell_text += ' ' + cell.text.lower()

                # Check for required data rows
                has_revenue = 'revenue' in all_cell_text
                has_expenses = 'expenses' in all_cell_text or 'expense' in all_cell_text
                has_net_profit = 'net profit' in all_cell_text or 'profit' in all_cell_text

                # Check for month columns
                has_july = 'july' in all_cell_text or 'jul' in all_cell_text
                has_august = 'august' in all_cell_text or 'aug' in all_cell_text
                has_september = 'september' in all_cell_text or 'sept' in all_cell_text or 'sep' in all_cell_text

                # Score table structure
                data_rows_found = sum([has_revenue, has_expenses, has_net_profit])
                months_found = sum([has_july, has_august, has_september])

                if data_rows_found == 3 and months_found == 3:
                    table_score = 0.25
                elif data_rows_found >= 2 and months_found >= 2:
                    table_score = 0.15
                elif data_rows_found >= 1 and months_found >= 1:
                    table_score = 0.08

                if table_score > best_table_score:
                    best_table_score = table_score

            if best_table_score >= 0.25:
                print(f"PASS: Component 4 — Financial table with Revenue/Expenses/Net Profit and Jul/Aug/Sept (0.25 pts)")
            elif best_table_score > 0:
                print(f"PARTIAL: Component 4 — Incomplete financial table ({best_table_score} pts)")
            else:
                print(f"FAIL: Component 4 — No table with required financial data found")
            total_score += best_table_score
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ---------- Component 5: Table of Contents section (0.10 points) ----------
    # Golden has a Heading 1 "Table of Contents".
    try:
        toc_found = False
        for s, t in all_paras:
            if 'table of contents' in t.lower() or 'toc' in t.lower():
                toc_found = True
                break

        if toc_found:
            print(f"PASS: Component 5 — Table of Contents section found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — No Table of Contents section found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # ---------- Component 6: Page numbers in footer (0.15 points) ----------
    # Golden has "Page " + PAGE field code in footer.
    try:
        footer_has_page_num = False
        for section in doc.sections:
            footer = section.footer
            if footer and footer.paragraphs:
                for fp in footer.paragraphs:
                    # Check for PAGE field code in XML
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    fld_chars = fp._element.findall('.//w:fldChar', ns)
                    instr_texts = fp._element.findall('.//w:instrText', ns)

                    for instr in instr_texts:
                        if instr.text and 'PAGE' in instr.text.upper():
                            footer_has_page_num = True
                            break

                    # Also check if footer text contains digits (cached page number)
                    if not footer_has_page_num and fp.text.strip():
                        if any(c.isdigit() for c in fp.text) or 'page' in fp.text.lower():
                            footer_has_page_num = True

                    if footer_has_page_num:
                        break
            if footer_has_page_num:
                break

        if footer_has_page_num:
            print(f"PASS: Component 6 — Page numbers found in footer (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 6 — No page numbers in footer")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
