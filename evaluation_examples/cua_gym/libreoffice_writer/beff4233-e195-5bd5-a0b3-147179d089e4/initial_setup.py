"""
Initial Setup: Create a Writer document with 10 bibliography entries cited in
the text body.  The bibliography section at the end is listed ALPHABETICALLY
(not by citation order) and uses author-year format — the agent's task is to
convert it to a numbered [1]...[10] list sorted by order of first citation.

Task ID: writer_bs_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_031'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# ---------- citation-order data (order of first appearance in text) ----------
# Index 0 → first cited, index 9 → last cited
CITATIONS_BY_ORDER = [
    {"author": "Baker",    "year": 2020, "title": "Scalable Microservice Architectures for Enterprise Systems",           "journal": "Journal of Software Engineering",           "vol": "45", "issue": "3", "pages": "112-130"},
    {"author": "Adams",    "year": 2019, "title": "A Survey of Natural Language Processing Techniques in Healthcare",     "journal": "Computational Linguistics Review",           "vol": "32", "issue": "1", "pages": "45-78"},
    {"author": "Clark",    "year": 2021, "title": "Deep Reinforcement Learning for Autonomous Vehicle Navigation",       "journal": "IEEE Transactions on Intelligent Systems",   "vol": "18", "issue": "7", "pages": "201-219"},
    {"author": "Davis",    "year": 2018, "title": "Statistical Methods for Large-Scale Genomic Data Analysis",            "journal": "Bioinformatics Advances",                    "vol": "9",  "issue": "2", "pages": "88-105"},
    {"author": "Evans",    "year": 2022, "title": "Quantum Computing Approaches to Combinatorial Optimization",           "journal": "Nature Computational Science",               "vol": "2",  "issue": "4", "pages": "310-325"},
    {"author": "Foster",   "year": 2020, "title": "Privacy-Preserving Machine Learning with Federated Architectures",     "journal": "ACM Computing Surveys",                      "vol": "53", "issue": "6", "pages": "1-35"},
    {"author": "Garcia",   "year": 2021, "title": "Edge Computing Paradigms for Internet of Things Applications",         "journal": "IEEE Internet of Things Journal",            "vol": "8",  "issue": "12","pages": "9750-9768"},
    {"author": "Harris",   "year": 2019, "title": "Transfer Learning Strategies for Low-Resource Language Models",        "journal": "Transactions of the ACL",                    "vol": "7",  "issue": "1", "pages": "154-172"},
    {"author": "Ibrahim",  "year": 2023, "title": "Explainable AI Methods for Clinical Decision Support Systems",        "journal": "Artificial Intelligence in Medicine",        "vol": "138","issue": "1", "pages": "1-18"},
    {"author": "Johnson",  "year": 2017, "title": "Graph Neural Networks: Foundations and Frontiers",                      "journal": "Machine Learning Research",                  "vol": "20", "issue": "5", "pages": "400-438"},
]

# ---------- alphabetical order (for the INITIAL bibliography) ----------
ALPHA_ORDER = sorted(range(len(CITATIONS_BY_ORDER)),
                     key=lambda i: CITATIONS_BY_ORDER[i]["author"])


def bib_entry_author_year(c):
    """Author-year format (initial state)."""
    return (f'{c["author"]} ({c["year"]}). "{c["title"]}." '
            f'{c["journal"]}, {c["vol"]}({c["issue"]}), pp. {c["pages"]}.')


def create_initial():
    doc = Document()

    # ---- styles (normal body) ----
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ---- Title ----
    title = doc.add_heading('Advances in Computational Intelligence: A Literature Review', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---- Abstract ----
    doc.add_heading('Abstract', level=2)
    doc.add_paragraph(
        'This paper provides a comprehensive review of recent advances across '
        'multiple sub-fields of computational intelligence, covering software '
        'engineering, natural language processing, autonomous systems, genomics, '
        'quantum computing, privacy-preserving methods, edge computing, transfer '
        'learning, explainable AI, and graph neural networks. We synthesize '
        'findings from ten influential studies published between 2017 and 2023.'
    )

    # ---- Introduction ----
    doc.add_heading('1. Introduction', level=2)
    doc.add_paragraph(
        'The rapid evolution of computational methods has transformed research '
        'and industry practices over the past decade. In this review we examine '
        'contributions that span architecture design (Baker, 2020), language '
        'understanding (Adams, 2019), and autonomous navigation (Clark, 2021). '
        'Earlier foundational work on genomic analysis (Davis, 2018) and graph '
        'neural networks (Johnson, 2017) laid the groundwork for more recent '
        'explorations into quantum optimization (Evans, 2022) and federated '
        'learning (Foster, 2020).'
    )

    # ---- Background ----
    doc.add_heading('2. Background', level=2)
    doc.add_paragraph(
        'Baker (2020) proposed a microservice architecture that decouples '
        'enterprise components, enabling horizontal scaling. Adams (2019) '
        'surveyed NLP techniques specifically tailored for clinical text, '
        'identifying key challenges in domain adaptation. Clark (2021) '
        'demonstrated that deep reinforcement learning agents can navigate '
        'complex urban environments with a 94% success rate under simulated '
        'traffic conditions.'
    )

    # ---- Methods ----
    doc.add_heading('3. Methodology', level=2)
    doc.add_paragraph(
        'Our review methodology follows the PRISMA guidelines. We searched '
        'five major databases for peer-reviewed publications between 2017 and '
        '2023. Davis (2018) provides the statistical framework we adopted for '
        'meta-analysis, while Evans (2022) informed our treatment of quantum '
        'speedup claims. The federated learning taxonomy from Foster (2020) '
        'structured our evaluation of privacy-aware systems.'
    )

    # ---- Results ----
    doc.add_heading('4. Results and Discussion', level=2)
    doc.add_paragraph(
        'Garcia (2021) showed that edge computing reduces inference latency '
        'by 40% for IoT sensor networks, complementing the low-resource '
        'transfer learning strategies of Harris (2019). Ibrahim (2023) '
        'recently extended explainable AI to multi-modal clinical data, '
        'achieving state-of-the-art interpretability scores. Johnson (2017) '
        'remains the most cited work in the graph neural network space, with '
        'over 3,200 citations as of 2024.'
    )

    # ---- Conclusion ----
    doc.add_heading('5. Conclusion', level=2)
    doc.add_paragraph(
        'The ten studies reviewed here represent complementary advances that '
        'collectively push the boundaries of computational intelligence. '
        'Future work should explore tighter integration of these methods, '
        'particularly combining federated learning (Foster, 2020) with '
        'explainable AI (Ibrahim, 2023) for trustworthy clinical systems.'
    )

    # ---- Bibliography (alphabetical, author-year format) ----
    doc.add_heading('References', level=2)
    for idx in ALPHA_ORDER:
        c = CITATIONS_BY_ORDER[idx]
        doc.add_paragraph(bib_entry_author_year(c))

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')


def launch_gui(command: str, delay_sec: float = 1.0):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


create_initial()
launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')
