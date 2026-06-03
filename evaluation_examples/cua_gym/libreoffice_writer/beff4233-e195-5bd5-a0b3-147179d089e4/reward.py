"""
Reward Script: Create formatted bibliography with numbered list [1]-[10] sorted by citation order
Task ID: writer_bs_031
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Bibliography entries have [N] numbered prefixes in correct order
  Component 2 (0.35): Bibliography entries are sorted by citation order (Baker first, Adams second, etc.)
  Component 3 (0.30): In-text citations use numbered [N] format instead of author-year
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_031'

# Expected citation order (order of first appearance in text)
# From task context: Baker(2020) first, Adams(2019) second, Clark(2021) third, etc.
EXPECTED_ORDER = [
    "Baker",
    "Adams",
    "Clark",
    "Davis",
    "Evans",
    "Foster",
    "Garcia",
    "Harris",
    "Ibrahim",
    "Johnson",
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate the References section
    # Find the paragraph index where "References" heading appears
    ref_start = None
    for i, p in enumerate(doc.paragraphs):
        if p.style and 'Heading' in p.style.name and 'References' in p.text:
            ref_start = i + 1
            break

    if ref_start is None:
        print("FAIL: No 'References' heading found")
        print("REWARD: 0.0")
        return 0.0

    # Collect bibliography entries (non-empty paragraphs after References heading)
    bib_entries = []
    for i in range(ref_start, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if text:
            bib_entries.append(text)

    print(f"INFO: Found {len(bib_entries)} bibliography entries")

    # Component 1: Bibliography entries have [N] numbered prefixes (0.35 points)
    # Each entry should start with [1], [2], ..., [10]
    try:
        numbered_count = 0
        correct_numbers = 0
        for idx, entry in enumerate(bib_entries):
            expected_num = idx + 1
            # Check if entry starts with [N]
            match = re.match(r'^\[(\d+)\]', entry)
            if match:
                numbered_count += 1
                actual_num = int(match.group(1))
                if actual_num == expected_num:
                    correct_numbers += 1

        if len(bib_entries) >= 10 and numbered_count >= 10 and correct_numbers >= 10:
            print(f"PASS: Component 1 — All {correct_numbers} entries have correct [N] numbering (0.35 pts)")
            total_score += 0.35
        elif numbered_count > 0:
            # Partial: some entries are numbered
            ratio = correct_numbers / max(len(bib_entries), 1)
            partial = round(0.35 * ratio, 2)
            print(f"PARTIAL: Component 1 — {correct_numbers}/{len(bib_entries)} entries correctly numbered ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No [N] numbered prefixes found in bibliography entries")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Bibliography entries are sorted by citation order with correct [N] prefix (0.35 points)
    # [1] Baker, [2] Adams, [3] Clark, etc. — must have BOTH the number prefix AND the correct author
    try:
        correct_order = 0
        for idx, entry in enumerate(bib_entries):
            if idx < len(EXPECTED_ORDER):
                expected_author = EXPECTED_ORDER[idx]
                expected_num = idx + 1
                # Must start with [N] AND contain the correct author
                match = re.match(r'^\[(\d+)\]', entry)
                if match and int(match.group(1)) == expected_num and expected_author in entry:
                    correct_order += 1
                else:
                    print(f"  Order mismatch at [{idx+1}]: expected '[{expected_num}] {expected_author}...', got '{entry[:60]}'")

        if correct_order >= 10:
            print(f"PASS: Component 2 — All 10 entries have correct [N] + author in citation order (0.35 pts)")
            total_score += 0.35
        elif correct_order > 0:
            ratio = correct_order / 10
            partial = round(0.35 * ratio, 2)
            print(f"PARTIAL: Component 2 — {correct_order}/10 entries correctly ordered ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Bibliography not in numbered citation order")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: In-text citations use numbered [N] format (0.30 points)
    # Check body paragraphs (before References) for [N] citation markers
    # and absence of author-year format like "Baker (2020)"
    try:
        # Collect body text (paragraphs before References heading)
        body_text = ""
        for i in range(0, ref_start):
            body_text += doc.paragraphs[i].text + "\n"

        # Check for numbered citations [1], [2], etc. in body
        numbered_citations = re.findall(r'\[(\d+)\]', body_text)
        has_numbered = len(numbered_citations) >= 5  # expect multiple numbered citations

        # Check that author-year citations are gone (e.g. "Baker (2020)" or "(Baker, 2020)")
        author_year_pattern = r'\b(?:Baker|Adams|Clark|Davis|Evans|Foster|Garcia|Harris|Ibrahim|Johnson)\s*\(\d{4}\)'
        author_year_matches = re.findall(author_year_pattern, body_text)
        no_author_year = len(author_year_matches) == 0

        if has_numbered and no_author_year:
            print(f"PASS: Component 3 — In-text citations use [N] format ({len(numbered_citations)} found), no author-year format (0.30 pts)")
            total_score += 0.30
        elif has_numbered:
            # Has numbered but still has some author-year leftovers
            print(f"PARTIAL: Component 3 — Has numbered citations but {len(author_year_matches)} author-year citations remain (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 — In-text citations not in [N] format (found {len(numbered_citations)} numbered, {len(author_year_matches)} author-year)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
