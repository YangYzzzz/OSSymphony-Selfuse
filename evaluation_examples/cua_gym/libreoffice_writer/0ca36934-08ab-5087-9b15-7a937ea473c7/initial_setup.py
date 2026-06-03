"""
Initial Setup: Scientific document with 'bioinformatics' and empty custom dictionary
Task ID: writer_fp_023
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_023'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

LO_CONFIG = '/home/user/.config/libreoffice/4/user'
WORDBOOK_DIR = f'{LO_CONFIG}/wordbook'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_document():
    """Create a scientific paper document that uses 'bioinformatics' multiple times."""
    doc = Document()

    # Title
    title = doc.add_heading('Advances in Computational Bioinformatics for Genomic Analysis', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Authors
    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = authors.add_run('Dr. Elena Vasquez, Prof. Takeshi Yamamoto, Dr. Priya Sharma')
    run.font.size = Pt(11)
    run.font.italic = True

    # Affiliation
    affil = doc.add_paragraph()
    affil.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = affil.add_run('Department of Bioinformatics and Computational Biology\n'
                        'Stanford University School of Medicine')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacing

    # Abstract
    doc.add_heading('Abstract', level=1)
    abstract = doc.add_paragraph(
        'The field of bioinformatics has undergone a transformative evolution in the past decade, '
        'driven by advances in high-throughput sequencing technologies and machine learning algorithms. '
        'This paper reviews the current state of bioinformatics methodologies applied to genomic analysis, '
        'with particular emphasis on next-generation sequencing data processing pipelines. We present '
        'a comprehensive survey of computational tools used in modern bioinformatics laboratories and '
        'discuss the challenges of scaling bioinformatics workflows to handle petabyte-scale datasets. '
        'Our findings suggest that the integration of deep learning with traditional bioinformatics '
        'approaches yields significant improvements in variant calling accuracy and structural '
        'variant detection.'
    )
    abstract.paragraph_format.space_after = Pt(12)

    # Keywords
    kw = doc.add_paragraph()
    run = kw.add_run('Keywords: ')
    run.bold = True
    kw.add_run('bioinformatics, genomics, next-generation sequencing, machine learning, '
               'variant calling, computational biology')

    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Bioinformatics represents a critical intersection of biology, computer science, and '
        'statistics. Since the completion of the Human Genome Project in 2003, the demand for '
        'sophisticated bioinformatics tools has grown exponentially. Modern bioinformatics '
        'encompasses a wide range of computational methods, from sequence alignment algorithms '
        'to complex network analysis of protein-protein interactions.'
    )
    doc.add_paragraph(
        'The rapid decrease in sequencing costs has made whole-genome sequencing accessible '
        'to clinical laboratories worldwide. This democratization of genomic data has placed '
        'new demands on the bioinformatics community to develop scalable, reproducible, and '
        'user-friendly analysis pipelines. Contemporary bioinformatics challenges include the '
        'interpretation of non-coding variants, the integration of multi-omics data, and the '
        'application of artificial intelligence to predict phenotypic outcomes from genotypic data.'
    )
    doc.add_paragraph(
        'In this review, we examine how bioinformatics has evolved from simple sequence '
        'comparison tools to sophisticated platforms that integrate genomic, transcriptomic, '
        'and proteomic data. We highlight key bioinformatics algorithms that have shaped the '
        'field and discuss emerging trends that will define the next generation of bioinformatics '
        'research.'
    )

    # 2. Methods in Computational Bioinformatics
    doc.add_heading('2. Methods in Computational Bioinformatics', level=1)
    doc.add_paragraph(
        'Modern bioinformatics relies on a diverse toolkit of computational methods. Sequence '
        'alignment remains the cornerstone of bioinformatics analysis, with tools such as BLAST, '
        'BWA, and Bowtie2 providing fast and accurate alignment of reads to reference genomes. '
        'The bioinformatics pipeline for variant discovery typically begins with quality control '
        'of raw sequencing reads, followed by alignment, duplicate marking, base quality score '
        'recalibration, and variant calling.'
    )

    # 2.1 Sequence Analysis
    doc.add_heading('2.1 Sequence Analysis', level=2)
    doc.add_paragraph(
        'Sequence analysis is perhaps the most fundamental bioinformatics operation. The Smith-Waterman '
        'algorithm for local alignment and the Needleman-Wunsch algorithm for global alignment remain '
        'foundational to bioinformatics education and practice. Hidden Markov Models (HMMs) have '
        'become integral to bioinformatics applications such as gene prediction, protein domain '
        'identification, and multiple sequence alignment refinement.'
    )

    # 2.2 Structural Bioinformatics
    doc.add_heading('2.2 Structural Bioinformatics', level=2)
    doc.add_paragraph(
        'Structural bioinformatics focuses on the three-dimensional structures of biological '
        'macromolecules. Recent breakthroughs, including AlphaFold2, have demonstrated that '
        'deep learning approaches can predict protein structures with near-experimental accuracy. '
        'This advancement has transformed structural bioinformatics from a niche specialty into '
        'a mainstream bioinformatics application with broad implications for drug discovery and '
        'enzyme engineering.'
    )

    # 3. Results and Discussion
    doc.add_heading('3. Results and Discussion', level=1)
    doc.add_paragraph(
        'Our survey of 287 bioinformatics laboratories across 42 countries reveals several '
        'key trends. First, Python has emerged as the dominant programming language in '
        'bioinformatics, used by 89% of respondents for daily analysis tasks. R remains '
        'important for statistical bioinformatics and visualization, used by 76% of labs. '
        'Second, cloud computing adoption in bioinformatics has increased from 23% in 2019 '
        'to 67% in 2025, reflecting the growing need for elastic computational resources '
        'in bioinformatics workflows.'
    )
    doc.add_paragraph(
        'The integration of machine learning into bioinformatics has accelerated dramatically. '
        'Neural network architectures, particularly transformer models adapted from natural '
        'language processing, have shown remarkable performance in bioinformatics tasks such as '
        'protein function prediction, drug-target interaction modeling, and single-cell RNA '
        'sequencing analysis. These bioinformatics applications of deep learning represent '
        'a paradigm shift from traditional feature-engineering approaches.'
    )

    # 4. Conclusion
    doc.add_heading('4. Conclusion', level=1)
    doc.add_paragraph(
        'Bioinformatics continues to evolve as new biological questions drive methodological '
        'innovation. The convergence of bioinformatics with artificial intelligence, cloud '
        'computing, and real-time clinical decision support systems promises to transform '
        'healthcare delivery. As the volume and complexity of biological data continue to '
        'grow, bioinformatics will play an increasingly central role in translating raw data '
        'into actionable biomedical knowledge. The future of bioinformatics lies in the '
        'development of interpretable, scalable, and clinically validated computational tools '
        'that bridge the gap between genomic discovery and patient care.'
    )

    # References
    doc.add_heading('References', level=1)
    refs = [
        'Altschul, S.F., et al. (1990). Basic local alignment search tool. J Mol Biol, 215(3), 403-410.',
        'Li, H. & Durbin, R. (2009). Fast and accurate short read alignment with Burrows-Wheeler transform. Bioinformatics, 25(14), 1754-1760.',
        'McKenna, A., et al. (2010). The Genome Analysis Toolkit: a MapReduce framework for analyzing NGS data. Genome Res, 20(9), 1297-1303.',
        'Jumper, J., et al. (2021). Highly accurate protein structure prediction with AlphaFold. Nature, 596(7873), 583-589.',
        'Vasquez, E. & Yamamoto, T. (2024). Deep learning frameworks for integrative bioinformatics. Nature Methods, 21(4), 412-425.',
        'Sharma, P., et al. (2025). Scalable bioinformatics pipelines for clinical genomics. Genome Biology, 26(1), 88.',
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f'[{i}] {ref}')
        p.paragraph_format.space_after = Pt(2)

    # Set page margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    doc.save(OUTPUT)
    print(f'Document created: {OUTPUT}')


def setup_custom_dictionary():
    """Create TechnicalTerms custom dictionary (empty, no hyphenation exceptions)."""
    os.makedirs(WORDBOOK_DIR, exist_ok=True)

    # Create TechnicalTerms.dic - a positive dictionary with no hyphenation entries
    dic_content = """OOoUserDict1
lang: <none>
type: 0
---
"""
    with open(f'{WORDBOOK_DIR}/TechnicalTerms.dic', 'w') as f:
        f.write(dic_content)
    print(f'Created TechnicalTerms.dic in {WORDBOOK_DIR}')

    # Register the dictionary in registrymodifications.xcu
    reg_file = f'{LO_CONFIG}/registrymodifications.xcu'

    # Read existing content
    if os.path.exists(reg_file):
        with open(reg_file, 'r') as f:
            content = f.read()
    else:
        content = '<?xml version="1.0" encoding="UTF-8"?>\n<oor:items xmlns:oor="http://openoffice.org/2001/registry" xmlns:xs="http://www.w3.org/2001/XMLSchema" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">\n</oor:items>'

    # Add dictionary list configuration and auto-hyphenation setting
    new_entries = '''<item oor:path="/org.openoffice.Office.Linguistic/ServiceManager/Dictionaries/HyphsAndSpells"><prop oor:name="Locations" oor:op="fuse"><value><it>%origin%/dictionaries/en/en_US.dic</it></value></prop></item>
<item oor:path="/org.openoffice.Office.Linguistic/ServiceManager/Dictionaries"><node oor:name="TechnicalTerms" oor:op="fuse"><prop oor:name="Locations"><value><it>%origin%/../../../../user/wordbook/TechnicalTerms.dic</it></value></prop><prop oor:name="Format"><value>7</value></prop><prop oor:name="Locales"><value><it/></value></prop></node></item>'''

    # Insert before closing tag
    if '</oor:items>' in content:
        content = content.replace('</oor:items>', new_entries + '\n</oor:items>')

    with open(reg_file, 'w') as f:
        f.write(content)
    print('Registered TechnicalTerms dictionary in registrymodifications.xcu')


def enable_auto_hyphenation():
    """Enable automatic hyphenation in LibreOffice Writer settings."""
    reg_file = f'{LO_CONFIG}/registrymodifications.xcu'
    with open(reg_file, 'r') as f:
        content = f.read()

    # Add auto-hyphenation settings
    hyphen_entries = '''<item oor:path="/org.openoffice.Office.Writer/Content/Format/Option"><prop oor:name="IsHyphAuto" oor:op="fuse"><value>true</value></prop></item>'''

    if '</oor:items>' in content:
        content = content.replace('</oor:items>', hyphen_entries + '\n</oor:items>')

    with open(reg_file, 'w') as f:
        f.write(content)
    print('Enabled auto-hyphenation in Writer settings')


def create_initial():
    create_document()
    setup_custom_dictionary()
    enable_auto_hyphenation()

    # Open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
