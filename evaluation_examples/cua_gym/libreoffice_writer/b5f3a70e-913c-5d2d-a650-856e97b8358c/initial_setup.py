"""
Initial Setup: Create a thesis document without widow/orphan control
Task ID: writer_acad_035
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_035'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default style font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # Explicitly disable widow/orphan control so single lines can appear alone
    # at page breaks (the task is to fix this)
    style.paragraph_format.widow_control = False

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title page
    title = doc.add_heading('The Impact of Artificial Intelligence on Modern Healthcare Systems: '
                            'A Comprehensive Analysis of Machine Learning Applications in Clinical Practice', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('\nA Thesis Submitted in Partial Fulfillment\n'
                           'of the Requirements for the Degree of\n'
                           'Master of Science in Computer Science\n\n'
                           'Department of Computer Science and Engineering\n'
                           'Stanford University\n\n'
                           'Elena Vasquez\n'
                           'March 2025')
    run.font.size = Pt(14)

    doc.add_page_break()

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This thesis investigates the transformative potential of artificial intelligence and machine '
        'learning technologies in modern healthcare delivery systems. Through a comprehensive analysis '
        'of 847 clinical deployments across 23 countries spanning from 2018 to 2024, we examine the '
        'efficacy, safety profiles, and implementation challenges associated with AI-assisted diagnostic '
        'tools, treatment planning algorithms, and patient monitoring frameworks. Our findings reveal '
        'that properly calibrated machine learning models can achieve diagnostic accuracy rates exceeding '
        '94.3% for certain categories of medical imaging analysis, while simultaneously reducing the '
        'average diagnostic turnaround time by approximately 67%. However, we also identify critical '
        'limitations related to algorithmic bias, data quality requirements, and the necessity for '
        'robust human oversight mechanisms. The research contributes a novel evaluation framework for '
        'assessing AI readiness in clinical environments and proposes a set of evidence-based guidelines '
        'for responsible deployment of these technologies in healthcare settings across diverse '
        'socioeconomic contexts.'
    )

    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        'Chapter 1: Introduction ...................................... 4',
        'Chapter 2: Literature Review ................................. 7',
        'Chapter 3: Methodology ...................................... 12',
        'Chapter 4: Results and Analysis ............................. 17',
        'Chapter 5: Discussion ....................................... 23',
        'Chapter 6: Conclusions and Future Work ...................... 28',
        'References .................................................. 31',
        'Appendices .................................................. 35',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # Chapter 1: Introduction
    doc.add_heading('Chapter 1: Introduction', level=1)

    doc.add_heading('1.1 Background and Motivation', level=2)
    doc.add_paragraph(
        'The intersection of artificial intelligence and healthcare represents one of the most '
        'promising frontiers in modern technology. Over the past decade, advances in computational '
        'power, data availability, and algorithmic sophistication have collectively enabled the '
        'development of AI systems capable of performing tasks that were previously the exclusive '
        'domain of trained medical professionals. From interpreting radiological images to predicting '
        'patient deterioration in intensive care units, these systems are fundamentally reshaping '
        'the landscape of clinical practice.'
    )
    doc.add_paragraph(
        'The global healthcare industry faces unprecedented challenges in the twenty-first century. '
        'An aging population, rising chronic disease burden, and persistent workforce shortages have '
        'created enormous pressure on healthcare delivery systems worldwide. According to the World '
        'Health Organization, the global shortage of healthcare workers is projected to reach 10 million '
        'by 2030, with the most severe deficits concentrated in low-income and lower-middle-income '
        'countries. In this context, AI-powered tools offer a compelling proposition: the potential to '
        'augment human capabilities, improve diagnostic efficiency, and extend quality healthcare '
        'services to underserved populations.'
    )
    doc.add_paragraph(
        'However, the deployment of artificial intelligence in clinical settings is not without '
        'significant challenges and concerns. Questions regarding algorithmic transparency, data '
        'privacy, regulatory compliance, and the potential for automation bias have generated '
        'substantial debate within both the medical and technology communities. Furthermore, the '
        'heterogeneity of healthcare systems across different countries and regions means that '
        'solutions developed in one context may not be directly transferable to another without '
        'careful adaptation and validation.'
    )

    doc.add_heading('1.2 Research Objectives', level=2)
    doc.add_paragraph(
        'This thesis aims to address several critical gaps in the existing literature on AI '
        'in healthcare. Specifically, this research pursues the following objectives:'
    )
    objectives = [
        'To systematically catalog and analyze the current landscape of AI deployments in clinical '
        'healthcare settings across diverse geographical and socioeconomic contexts.',
        'To evaluate the diagnostic performance, safety profiles, and implementation outcomes of '
        'machine learning models deployed in real-world clinical environments.',
        'To identify the key barriers and facilitating factors that influence the successful adoption '
        'of AI technologies in healthcare organizations of varying sizes and resource levels.',
        'To develop a comprehensive evaluation framework that can guide healthcare institutions in '
        'assessing their readiness for AI integration.',
        'To propose evidence-based guidelines for the responsible and equitable deployment of AI '
        'technologies across healthcare systems globally.',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Number')

    doc.add_heading('1.3 Scope and Limitations', level=2)
    doc.add_paragraph(
        'The scope of this research encompasses AI applications in clinical healthcare delivery, '
        'with a particular focus on diagnostic imaging, clinical decision support systems, and '
        'patient monitoring applications. Administrative uses of AI in healthcare, such as billing '
        'optimization and scheduling, fall outside the scope of this study. Similarly, while we '
        'acknowledge the growing importance of AI in pharmaceutical research and drug discovery, '
        'these applications are not included in our analysis due to the fundamentally different '
        'regulatory and validation requirements they entail.'
    )
    doc.add_paragraph(
        'Our analysis is limited to AI deployments that have progressed beyond the pilot or '
        'proof-of-concept stage and have been integrated into routine clinical workflows. This '
        'criterion was established to ensure that our findings reflect the practical realities of '
        'AI implementation rather than theoretical capabilities demonstrated only in controlled '
        'research environments. Additionally, our data collection was restricted to deployments '
        'documented in peer-reviewed literature, institutional reports, or verified through direct '
        'communication with deploying organizations.'
    )

    # Chapter 2: Literature Review
    doc.add_heading('Chapter 2: Literature Review', level=1)

    doc.add_heading('2.1 Historical Development of Medical AI', level=2)
    doc.add_paragraph(
        'The application of computational methods to medical diagnosis has a history spanning over '
        'five decades. Early expert systems such as MYCIN, developed at Stanford University in the '
        'mid-1970s, demonstrated that rule-based computer programs could match or exceed the '
        'performance of human specialists in narrowly defined diagnostic tasks. MYCIN achieved '
        'approximately 69% accuracy in recommending appropriate antibiotic therapies for blood '
        'infections, outperforming the 42.5% to 62.5% accuracy rates observed among infectious '
        'disease specialists consulted during the same evaluation period.'
    )
    doc.add_paragraph(
        'Despite these promising early results, the widespread adoption of clinical AI systems '
        'remained elusive for several decades. The limitations of rule-based systems, including '
        'their inability to handle uncertainty and their requirement for exhaustive knowledge '
        'engineering, constrained their practical utility in the complex and dynamic environment '
        'of clinical practice. The resurgence of interest in medical AI can be traced to the '
        'convergence of three critical developments in the early 2010s: the availability of large '
        'digitized medical datasets through electronic health record adoption, dramatic improvements '
        'in computing hardware particularly through graphics processing units, and theoretical '
        'breakthroughs in deep learning architectures.'
    )

    doc.add_heading('2.2 Machine Learning in Diagnostic Imaging', level=2)
    doc.add_paragraph(
        'Diagnostic imaging has emerged as the most active area of clinical AI research and '
        'deployment. The visual nature of imaging data makes it particularly amenable to analysis '
        'by convolutional neural networks, which have demonstrated remarkable capabilities in '
        'pattern recognition tasks. In 2016, Esteva and colleagues published a landmark study '
        'demonstrating that a deep neural network trained on a dataset of 129,450 clinical images '
        'could classify skin cancer with accuracy comparable to that of board-certified dermatologists.'
    )
    doc.add_paragraph(
        'Subsequent studies have extended these findings to a wide range of imaging modalities and '
        'clinical applications. In ophthalmology, AI systems have been developed for the automated '
        'detection of diabetic retinopathy, age-related macular degeneration, and glaucoma from '
        'retinal fundus photographs. The FDA-approved IDx-DR system, for example, demonstrated '
        'sensitivity of 87.2% and specificity of 90.7% for detecting more-than-mild diabetic '
        'retinopathy in a prospective clinical trial. In radiology, deep learning models have '
        'achieved promising results in detecting pulmonary nodules on chest CT scans, identifying '
        'breast cancer on mammograms, and detecting fractures on skeletal radiographs.'
    )

    doc.add_heading('2.3 Clinical Decision Support Systems', level=2)
    doc.add_paragraph(
        'Beyond diagnostic imaging, machine learning algorithms are increasingly being deployed '
        'in clinical decision support systems that analyze structured and unstructured patient '
        'data to assist clinicians in treatment planning and risk assessment. Sepsis prediction '
        'algorithms, which analyze vital signs, laboratory results, and nursing documentation to '
        'identify patients at risk of developing life-threatening infections, represent one of the '
        'most widely deployed categories of clinical AI. The implementation of such systems has '
        'been associated with reductions in sepsis mortality ranging from 12% to 26% in published '
        'studies, although the quality and rigor of this evidence varies considerably.'
    )
    doc.add_paragraph(
        'Natural language processing technologies are also gaining traction in clinical applications. '
        'These systems can extract clinically relevant information from unstructured clinical notes, '
        'pathology reports, and discharge summaries, enabling more comprehensive analysis of patient '
        'histories and facilitating clinical research through automated cohort identification. Recent '
        'advances in large language models have further expanded the potential applications of NLP '
        'in healthcare, including automated clinical documentation, patient communication, and '
        'literature synthesis for evidence-based practice.'
    )

    # Chapter 3: Methodology
    doc.add_heading('Chapter 3: Methodology', level=1)

    doc.add_heading('3.1 Research Design', level=2)
    doc.add_paragraph(
        'This study employs a mixed-methods research design combining systematic review '
        'methodology with quantitative meta-analysis and qualitative case study analysis. '
        'The research was conducted in three phases over a period of eighteen months, from '
        'January 2023 through June 2024. The systematic review component follows the PRISMA '
        'guidelines for reporting systematic reviews and meta-analyses, while the qualitative '
        'component adheres to the COREQ checklist for qualitative research reporting.'
    )

    doc.add_heading('3.2 Data Collection', level=2)
    doc.add_paragraph(
        'The primary data collection involved a comprehensive search of six electronic databases: '
        'PubMed, Embase, Cochrane Library, IEEE Xplore, ACM Digital Library, and Web of Science. '
        'Search terms were developed iteratively in consultation with a medical librarian and '
        'included combinations of terms related to artificial intelligence, machine learning, deep '
        'learning, neural networks, healthcare, clinical, diagnostic, and treatment. The search '
        'was restricted to publications in English from January 2018 through December 2023.'
    )
    doc.add_paragraph(
        'The initial search yielded 12,847 publications, which were screened through a two-stage '
        'process. Title and abstract screening, performed independently by two researchers, reduced '
        'the pool to 2,134 potentially eligible studies. Full-text review of these publications '
        'against our predetermined inclusion and exclusion criteria resulted in a final sample of '
        '847 studies documenting clinical AI deployments across 23 countries. Inter-rater reliability '
        'for study inclusion decisions, measured by Cohen kappa, was 0.89, indicating strong agreement.'
    )

    doc.add_heading('3.3 Analytical Framework', level=2)
    doc.add_paragraph(
        'Quantitative data from the included studies were analyzed using random-effects meta-analysis '
        'models implemented in R using the metafor package. Heterogeneity was assessed using the I '
        'squared statistic and Cochran Q test. Publication bias was evaluated through visual '
        'inspection of funnel plots and formal testing using the Egger regression method. Subgroup '
        'analyses were conducted by geographic region, income level of the deploying country, clinical '
        'domain, and type of AI algorithm employed.'
    )
    doc.add_paragraph(
        'The qualitative component consisted of semi-structured interviews with 45 key informants '
        'from 28 healthcare organizations across 12 countries. Informants included clinical '
        'informaticists, department heads who had championed AI adoption, front-line clinicians '
        'who used AI tools in their daily practice, and hospital administrators responsible for '
        'technology procurement and governance. Interviews were conducted via video conference, '
        'lasted between 45 and 90 minutes, and were recorded and transcribed verbatim. Thematic '
        'analysis was performed using a hybrid deductive-inductive approach guided by the '
        'Consolidated Framework for Implementation Research.'
    )

    # Chapter 4: Results
    doc.add_heading('Chapter 4: Results and Analysis', level=1)

    doc.add_heading('4.1 Overview of AI Deployments', level=2)
    doc.add_paragraph(
        'The 847 clinical AI deployments identified in our review span 23 countries across six '
        'continents. The geographic distribution reveals significant concentration, with 62.3% of '
        'all documented deployments located in just five countries: the United States (28.7%), '
        'China (14.2%), the United Kingdom (8.9%), Germany (5.8%), and South Korea (4.7%). '
        'Low-income countries collectively account for only 3.4% of documented deployments, '
        'highlighting a substantial equity gap in access to clinical AI technologies.'
    )
    doc.add_paragraph(
        'Diagnostic imaging applications comprise the largest category at 41.2% of all deployments, '
        'followed by clinical decision support systems at 23.8%, patient monitoring and early '
        'warning systems at 18.1%, and other applications including surgical robotics and pathology '
        'at 16.9%. The predominance of imaging applications is consistent with the technical '
        'maturity of computer vision algorithms and the relatively well-defined nature of many '
        'imaging interpretation tasks.'
    )

    doc.add_heading('4.2 Performance Metrics', level=2)
    doc.add_paragraph(
        'Meta-analysis of diagnostic accuracy across imaging applications yielded a pooled '
        'sensitivity of 91.7% with a 95% confidence interval of 89.2% to 93.8%, and a pooled '
        'specificity of 93.4% with a 95% confidence interval of 91.1% to 95.2%. These results '
        'indicate that AI imaging tools generally achieve performance comparable to that of '
        'specialist physicians, though significant heterogeneity was observed across clinical '
        'domains. The I squared statistic for the overall sensitivity estimate was 87.3%, '
        'indicating substantial between-study heterogeneity.'
    )
    doc.add_paragraph(
        'Subgroup analysis revealed notable variations in performance across different clinical '
        'applications. Dermatological image classification demonstrated the highest pooled accuracy '
        'at 94.3%, while chest radiograph interpretation showed more modest performance at 87.6%. '
        'Ophthalmological applications, particularly diabetic retinopathy screening, achieved '
        'consistently high performance with a pooled sensitivity of 92.8% and specificity of 91.3%. '
        'These differences likely reflect variations in task complexity, data availability for '
        'model training, and the inherent difficulty of the underlying clinical classification tasks.'
    )

    # Chapter 5: Discussion
    doc.add_heading('Chapter 5: Discussion', level=1)

    doc.add_heading('5.1 Key Findings and Implications', level=2)
    doc.add_paragraph(
        'The findings of this research paint a nuanced picture of the current state of AI in '
        'clinical healthcare. While the technology has demonstrated impressive capabilities in '
        'controlled evaluations and is increasingly being deployed in real-world clinical settings, '
        'significant challenges remain in ensuring equitable access, maintaining clinical safety, '
        'and achieving sustainable integration into existing healthcare workflows. The concentration '
        'of deployments in high-income countries and the persistent challenges around algorithmic '
        'bias raise important questions about whether AI in healthcare risks exacerbating rather '
        'than mitigating existing health inequities.'
    )
    doc.add_paragraph(
        'Our analysis reveals that the gap between AI performance in research settings and '
        'real-world clinical practice is often substantial. Several studies in our review reported '
        'significant decreases in diagnostic accuracy when models trained on data from one '
        'institution were deployed at another, underscoring the importance of local validation '
        'and continuous monitoring. This finding has profound implications for the scalability of '
        'clinical AI solutions and suggests that the development of robust, generalizable models '
        'remains an important area for future research.'
    )

    doc.add_heading('5.2 Barriers to Adoption', level=2)
    doc.add_paragraph(
        'The qualitative analysis identified several recurring themes related to barriers to AI '
        'adoption in healthcare. Trust emerged as the single most frequently cited barrier, '
        'mentioned by 38 of 45 interview participants. Clinicians expressed concerns about the '
        'opacity of AI decision-making processes and the potential for errors in edge cases that '
        'fall outside the training distribution. Several participants noted that their willingness '
        'to rely on AI recommendations was strongly influenced by their ability to understand the '
        'reasoning behind specific predictions.'
    )
    doc.add_paragraph(
        'Regulatory uncertainty was identified as the second most significant barrier, cited by '
        '33 participants. The rapidly evolving regulatory landscape for medical AI, with different '
        'approaches being taken by the FDA in the United States, the CE marking process in Europe, '
        'and regulatory bodies in Asia, creates complexity for organizations operating in multiple '
        'jurisdictions. Several participants from smaller healthcare organizations noted that the '
        'cost and complexity of regulatory compliance for AI systems represented a significant '
        'deterrent to adoption, particularly when the perceived clinical benefit was incremental '
        'rather than transformative.'
    )

    # Chapter 6: Conclusions
    doc.add_heading('Chapter 6: Conclusions and Future Work', level=1)

    doc.add_heading('6.1 Summary of Contributions', level=2)
    doc.add_paragraph(
        'This thesis has presented a comprehensive analysis of the current landscape of AI in '
        'clinical healthcare, drawing on evidence from 847 deployments across 23 countries. '
        'The research makes several contributions to the field. First, it provides the most '
        'extensive systematic mapping of clinical AI deployments to date, revealing both the '
        'breadth of current applications and the significant geographic and socioeconomic '
        'disparities in access to these technologies. Second, the meta-analytic findings offer '
        'robust estimates of diagnostic performance across major clinical domains, providing a '
        'reliable benchmark for future development and evaluation efforts.'
    )
    doc.add_paragraph(
        'Third, the qualitative insights from key informants across diverse healthcare settings '
        'illuminate the complex interplay of organizational, technical, and human factors that '
        'determine the success or failure of clinical AI implementations. The evaluation framework '
        'and deployment guidelines proposed in this thesis synthesize these findings into practical '
        'tools that can support healthcare organizations in navigating the complexities of AI '
        'adoption. We hope that these contributions will facilitate more informed, equitable, and '
        'effective integration of AI technologies into healthcare systems worldwide.'
    )

    doc.add_heading('6.2 Future Research Directions', level=2)
    doc.add_paragraph(
        'Several avenues for future research emerge from this work. Longitudinal studies tracking '
        'the evolution of AI system performance over extended deployment periods are critically '
        'needed to understand issues such as data drift, model degradation, and the long-term '
        'impact on clinical workflows and patient outcomes. Additionally, research examining the '
        'economic impact of clinical AI deployments, including comprehensive cost-effectiveness '
        'analyses and studies of workforce implications, would provide valuable evidence for '
        'healthcare policy makers and administrators.'
    )
    doc.add_paragraph(
        'The development of methods for improving the generalizability and fairness of clinical '
        'AI models across diverse patient populations represents another important research '
        'priority. Federated learning approaches, which enable model training on distributed '
        'datasets without centralizing sensitive patient data, offer a promising avenue for '
        'addressing both the generalizability and privacy challenges that currently constrain the '
        'development of robust clinical AI systems. Finally, research on effective strategies for '
        'human-AI collaboration in clinical settings, including studies of optimal interface design '
        'and decision support workflows, would help ensure that AI tools are integrated in ways '
        'that genuinely augment rather than undermine clinical expertise and judgment.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
