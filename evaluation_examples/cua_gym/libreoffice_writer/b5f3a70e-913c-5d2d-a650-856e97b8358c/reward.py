"""
Reward Script: Widow and orphan control for thesis
Task ID: writer_acad_035
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): Normal/Default Paragraph Style has widow_control enabled (True)
  Component 2 (0.4): Effective widow control for ALL paragraphs is True
                      (style-level True + no paragraph-level False overrides)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_035'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify widow and orphan control is properly configured.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: file must have paragraphs (sanity check)
    if len(doc.paragraphs) < 5:
        print(f"PRECONDITION FAIL: Document has only {len(doc.paragraphs)} paragraphs, expected a thesis document")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Normal style has widow_control enabled (0.6 points)
    # In the initial file, Normal style has widow_control=False (w:val="0").
    # The task requires enabling it (widow_control=True).
    try:
        normal_style = doc.styles['Normal']
        wc_value = normal_style.paragraph_format.widow_control
        if wc_value is True:
            print(f"PASS: Component 1 -- Normal style widow_control is True (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 1 -- Normal style widow_control is {wc_value}, expected True")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Effective widow control for ALL paragraphs is True (0.4 points)
    # This checks that every paragraph either:
    #   - Inherits from the Normal style (paragraph_format.widow_control is None) and style is True, OR
    #   - Has explicit widow_control=True
    # A paragraph with explicit widow_control=False would undermine the setting.
    # This component only passes if Component 1 also passes (style must be True for inheritance to work).
    try:
        normal_style = doc.styles['Normal']
        style_wc = normal_style.paragraph_format.widow_control

        disabled_count = 0
        total_paras = len(doc.paragraphs)

        for i, para in enumerate(doc.paragraphs):
            para_wc = para.paragraph_format.widow_control
            # Effective value: if para_wc is None, inherit from style; otherwise use para_wc
            if para_wc is None:
                effective_wc = style_wc
            else:
                effective_wc = para_wc

            if effective_wc is not True:
                disabled_count += 1

        if disabled_count == 0:
            print(f"PASS: Component 2 -- All {total_paras} paragraphs have effective widow_control=True (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 -- {disabled_count}/{total_paras} paragraphs have effective widow_control != True")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist app state before verification
persist_app_state("libreoffice_writer")

# Test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
