"""
Initial Setup: Short story manuscript in default formatting (pre-task state)
Task ID: writer_creative_004
Domain: libreoffice_writer

Creates lighthouse_keeper.docx on the Desktop with:
- 11pt Arial, single-spaced, 0.79-inch margins
- Title left-aligned, no header, no first-line indent
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'lighthouse_keeper'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default (narrow) margins — 0.79 inch ≈ 2.0 cm
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Story paragraphs content
    story_paragraphs = [
        "The old lighthouse stood at the edge of Widow's Point, its white paint peeling like sunburned skin. "
        "Thomas Alcott had tended the light for thirty-seven years, long enough to know every creak of the "
        "iron staircase, every moan of the wind through the gallery rails. He climbed the two hundred and "
        "twelve steps each evening at dusk, the number fixed in his bones.",

        "His daughter Elena had left the island four years ago, taking with her only a duffel bag and a "
        "photograph of the two of them standing at the lantern room window. The ferry had been swallowed "
        "by morning fog before Thomas reached the end of the dock. He had not called after her. There was "
        "nothing left to say that the silence hadn't already said better.",

        "October brought rough weather from the northeast, and Thomas spent the first three days of the "
        "month patching the seams on the fog horn housing. The work kept his hands occupied and his mind "
        "mercifully blank. When the storms passed, the sea turned to hammered pewter, and the container "
        "ships moved like slow ideas across the horizon.",

        "On the fourth day a supply boat arrived with provisions and a letter addressed in a handwriting "
        "Thomas didn't recognize. He set it on the kitchen table and cooked his supper first — rice and "
        "salted cod — then washed his bowl and dried it before finally sitting down with the envelope. "
        "The postmark read Portland, Maine.",

        "Inside was a single sheet of cream paper folded in thirds. The letter was from a woman named "
        "Clara Voss, who identified herself as Elena's roommate. She wrote in a careful, looping script "
        "that Elena had been in an accident in August. She was recovering, Clara explained, but had "
        "developed a fear of phones since the hospital. Elena had asked her to write on her behalf.",

        "Thomas read the letter twice, then set it face down on the table. Outside, a gull called once "
        "and fell silent. He thought about the photograph on Elena's wall — the two of them squinting "
        "against the lantern's glow, neither smiling but standing close, their shoulders almost touching. "
        "He had never been a man who could say the things he meant.",

        "That night he climbed the stairs more slowly than usual, pausing at the midpoint landing to look "
        "through the small round window at the black water below. The light made its steady revolution, "
        "indifferent and faithful, throwing its beam across the rocks and the dark sea beyond. Ships knew "
        "where they were because of it. Thomas had always taken a quiet pride in that.",

        "He sat down at the small desk in the watch room, pulled out a sheet of paper from the logbook "
        "drawer, and uncapped his pen. He had no idea what to say. He wrote her name at the top of the "
        "page, then crossed it out, then wrote it again. The light turned. The sea moved beneath the "
        "station. He began, at last, with the only honest thing he could think of: I should have called.",
    ]

    # Helper: set all runs in a paragraph to Arial 11pt
    def apply_default_font(para):
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    # Line 1: Title (left-aligned)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_para.paragraph_format.line_spacing = 1.0
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(0)
    run = title_para.add_run('The Lighthouse Keeper')
    run.font.name = 'Arial'
    run.font.size = Pt(11)

    # Line 2: Byline (left-aligned)
    byline_para = doc.add_paragraph()
    byline_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    byline_para.paragraph_format.line_spacing = 1.0
    byline_para.paragraph_format.space_before = Pt(0)
    byline_para.paragraph_format.space_after = Pt(0)
    run = byline_para.add_run('by Sarah J. Mitchell')
    run.font.name = 'Arial'
    run.font.size = Pt(11)

    # Story paragraphs (8 paragraphs, no indent)
    for text in story_paragraphs:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.first_line_indent = Pt(0)
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
