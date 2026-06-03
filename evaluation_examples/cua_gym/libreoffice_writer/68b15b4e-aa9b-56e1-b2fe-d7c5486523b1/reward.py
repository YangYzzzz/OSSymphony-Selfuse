"""
Reward Script: Standard submission format for short story manuscript
Task ID: writer_creative_004
Domain: libreoffice_writer
Scoring:
  Component 1: Font family = Courier New, size = 12pt (0.25 pts)
  Component 2: Line spacing = double (2.0) for all paragraphs (0.20 pts)
  Component 3: All margins = 1 inch (0.15 pts)
  Component 4: Title and byline centered (0.15 pts)
  Component 5: Header right-aligned with "Mitchell" + page number, first page excluded (0.15 pts)
  Component 6: First-line indent = 0.5 inch on story body paragraphs (0.10 pts)
  Total: 1.00
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_004'
FILE_PATH = '/home/user/Desktop/lighthouse_keeper.docx'

# Tolerance constants
MARGIN_TOLERANCE_EMU = 18000    # ~0.02 inch tolerance for margin checks
FONT_SIZE_TOLERANCE  = 0.5      # 0.5 pt tolerance


def _emu_to_inches(emu):
    return emu / 914400.0


def _emu_to_pt(emu):
    return emu / 12700.0


def verify_task(file_path):
    """
    Verify manuscript formatting task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # --- Precondition gate ---
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if not doc.paragraphs:
        print("CRITICAL: Document has no paragraphs.")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: Font = Courier New, 12pt for all non-empty paragraphs (0.25 pts)
    # -----------------------------------------------------------------------
    try:
        courier_ok = True
        size_ok = True
        checked_runs = 0

        for para in doc.paragraphs:
            for run in para.runs:
                if not run.text.strip():
                    continue
                checked_runs += 1
                # Font name check
                if run.font.name is not None and run.font.name != "Courier New":
                    courier_ok = False
                # Font size check (152400 EMU = 12pt)
                if run.font.size is not None:
                    size_pt = _emu_to_pt(run.font.size)
                    if abs(size_pt - 12.0) > FONT_SIZE_TOLERANCE:
                        size_ok = False

        if checked_runs == 0:
            print("FAIL: Component 1 — no non-empty runs found to check")
        elif courier_ok and size_ok:
            print(f"PASS: Component 1 — Courier New 12pt verified across {checked_runs} runs (0.25 pts)")
            total_score += 0.25
        else:
            if not courier_ok:
                print(f"FAIL: Component 1 — font name is not Courier New in at least one run")
            if not size_ok:
                print(f"FAIL: Component 1 — font size is not 12pt in at least one run")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: Line spacing = double (2.0) for all paragraphs (0.20 pts)
    # -----------------------------------------------------------------------
    try:
        spacing_ok = True
        checked_paras = 0

        for para in doc.paragraphs:
            # Skip completely empty paragraphs (spacing may be 0/None)
            if not para.text.strip():
                continue
            checked_paras += 1
            pf = para.paragraph_format
            ls = pf.line_spacing
            if ls is None:
                # Inheriting from style — check the style
                style_ls = para.style.paragraph_format.line_spacing if para.style else None
                if style_ls is None or abs(float(style_ls) - 2.0) > 0.05:
                    spacing_ok = False
            else:
                # line_spacing can be a float (multiple) or an EMU (absolute)
                # For double spacing it should be 2.0 (float/multiple)
                try:
                    ls_float = float(ls)
                    if abs(ls_float - 2.0) > 0.05:
                        spacing_ok = False
                except (TypeError, ValueError):
                    spacing_ok = False

        if checked_paras == 0:
            print("FAIL: Component 2 — no non-empty paragraphs found")
        elif spacing_ok:
            print(f"PASS: Component 2 — double line spacing (2.0) verified across {checked_paras} paragraphs (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — line spacing is not 2.0 in at least one paragraph")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: All margins = 1 inch (914400 EMU) (0.15 pts)
    # -----------------------------------------------------------------------
    try:
        TARGET_MARGIN_EMU = 914400  # 1 inch
        all_margins_ok = True
        section = doc.sections[0]
        margins = {
            'top': section.top_margin,
            'bottom': section.bottom_margin,
            'left': section.left_margin,
            'right': section.right_margin,
        }
        for side, val in margins.items():
            if val is None or abs(val - TARGET_MARGIN_EMU) > MARGIN_TOLERANCE_EMU:
                all_margins_ok = False
                print(f"FAIL: Component 3 — {side} margin = {_emu_to_inches(val):.3f} inches, expected 1.0 inch")

        if all_margins_ok:
            print("PASS: Component 3 — all margins are 1 inch (0.15 pts)")
            total_score += 0.15
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: Title and byline are centered (0.15 pts)
    # -----------------------------------------------------------------------
    try:
        title_centered = False
        byline_centered = False
        title_text_found = False
        byline_text_found = False

        for para in doc.paragraphs:
            text_lower = para.text.strip().lower()
            if "the lighthouse keeper" in text_lower:
                title_text_found = True
                if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    title_centered = True
                else:
                    print(f"FAIL: Component 4 — title alignment = {para.paragraph_format.alignment}, expected CENTER")
            elif "sarah j. mitchell" in text_lower or ("by" in text_lower and "mitchell" in text_lower):
                byline_text_found = True
                if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    byline_centered = True
                else:
                    print(f"FAIL: Component 4 — byline alignment = {para.paragraph_format.alignment}, expected CENTER")

        if not title_text_found:
            print("FAIL: Component 4 — title text 'The Lighthouse Keeper' not found")
        if not byline_text_found:
            print("FAIL: Component 4 — byline text not found")

        if title_centered and byline_centered:
            print("PASS: Component 4 — title and byline are both centered (0.15 pts)")
            total_score += 0.15
        elif title_centered or byline_centered:
            print(f"PARTIAL: Component 4 — only one of title/byline is centered")
        else:
            if title_text_found or byline_text_found:
                print("FAIL: Component 4 — neither title nor byline is centered")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -----------------------------------------------------------------------
    # Component 5: Header on pages 2+ has "Mitchell" + page number, right-aligned;
    #              first page header is empty/suppressed (0.15 pts)
    # -----------------------------------------------------------------------
    try:
        section = doc.sections[0]
        header_ok = False
        first_page_suppressed = False

        # Check if different first-page header is enabled
        different_first = section.different_first_page_header_footer
        if different_first:
            # First page header should be empty
            fp_header = section.first_page_header
            fp_text = "".join(p.text for p in fp_header.paragraphs).strip()
            if fp_text == "":
                first_page_suppressed = True
            else:
                print(f"FAIL: Component 5 — first page header should be empty, found: {repr(fp_text)}")
        else:
            # If no different first page, check if there's no content / standard header
            # Some implementations just leave header empty or use titlePg attribute differently
            first_page_suppressed = False
            print("FAIL: Component 5 — different_first_page_header_footer is not set (first page header not suppressed)")

        # Check the main header
        header = section.header
        header_paras = header.paragraphs
        if header_paras:
            main_header_para = header_paras[0]
            header_text = main_header_para.text.strip()
            header_align = main_header_para.paragraph_format.alignment

            # Header should contain "Mitchell" and a page number field code
            has_mitchell = "Mitchell" in header_text
            # Check for PAGE field code in XML
            header_xml = main_header_para._element.xml
            has_page_field = "PAGE" in header_xml or "fldChar" in header_xml

            is_right_aligned = (header_align == WD_PARAGRAPH_ALIGNMENT.RIGHT)

            if has_mitchell and has_page_field and is_right_aligned:
                header_ok = True
                print(f"PASS: Component 5 — main header has 'Mitchell', page field, right-aligned: {repr(header_text)}")
            else:
                issues = []
                if not has_mitchell:
                    issues.append(f"missing 'Mitchell' (found: {repr(header_text)})")
                if not has_page_field:
                    issues.append("missing PAGE field code")
                if not is_right_aligned:
                    issues.append(f"alignment={header_align}, expected RIGHT")
                print(f"FAIL: Component 5 — header issues: {', '.join(issues)}")
        else:
            print("FAIL: Component 5 — main header has no paragraphs")

        if header_ok and first_page_suppressed:
            print("PASS: Component 5 — header and first-page suppression correct (0.15 pts)")
            total_score += 0.15
        elif header_ok:
            print("PARTIAL: Component 5 — header content correct but first page not suppressed")
        elif first_page_suppressed:
            print("PARTIAL: Component 5 — first page suppressed but header content incorrect")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -----------------------------------------------------------------------
    # Component 6: Story body paragraphs have 0.5-inch first-line indent (0.10 pts)
    # -----------------------------------------------------------------------
    try:
        TARGET_INDENT_EMU = int(0.5 * 914400)  # 457200 EMU = 0.5 inches
        INDENT_TOLERANCE_EMU = 18000  # ~0.02 inch

        body_paras_checked = 0
        body_paras_indented = 0

        for para in doc.paragraphs:
            text_strip = para.text.strip()
            # Skip title, byline, and empty paragraphs — only check story body
            if not text_strip:
                continue
            if "the lighthouse keeper" in text_strip.lower():
                continue
            if "sarah j. mitchell" in text_strip.lower() or (text_strip.lower().startswith("by ") and "mitchell" in text_strip.lower()):
                continue
            body_paras_checked += 1
            pf = para.paragraph_format
            fli = pf.first_line_indent
            if fli is not None and abs(fli - TARGET_INDENT_EMU) <= INDENT_TOLERANCE_EMU:
                body_paras_indented += 1

        if body_paras_checked == 0:
            print("FAIL: Component 6 — no story body paragraphs found")
        elif body_paras_indented == body_paras_checked:
            print(f"PASS: Component 6 — all {body_paras_checked} body paragraphs have 0.5-inch first-line indent (0.10 pts)")
            total_score += 0.10
        elif body_paras_indented > 0:
            ratio = body_paras_indented / body_paras_checked
            print(f"PARTIAL: Component 6 — {body_paras_indented}/{body_paras_checked} body paragraphs have 0.5-inch indent")
        else:
            print(f"FAIL: Component 6 — none of {body_paras_checked} body paragraphs have 0.5-inch first-line indent")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


# --- Entrypoint ---
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
