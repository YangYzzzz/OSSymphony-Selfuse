"""
Initial Setup: Insert two new rows after row 3 in shipping manifest table
Task ID: writer_tm_028
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_028'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("Shipping Manifest", level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle info
    doc.add_paragraph("Vessel: MV Pacific Pioneer | Voyage: VP-2026-0041 | Date: 2026-03-28")

    # Create 5-column x 6-row table (1 header + 5 data rows)
    table = doc.add_table(rows=6, cols=5)
    table.style = "Table Grid"

    # Row 1: Headers
    headers = ["Item No", "Description", "Quantity", "Weight (kg)", "Destination"]
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    # Rows 2-6: Realistic shipping data
    data = [
        ["SM-1001", "Industrial Bearings (SKF 6205)", "240", "1,080.00", "Rotterdam, NL"],
        ["SM-1002", "Copper Wire Spools (2.5mm)", "85", "3,400.00", "Hamburg, DE"],
        ["SM-1003", "Automotive Brake Pads (Set)", "500", "750.00", "Felixstowe, UK"],
        ["SM-1004", "Stainless Steel Pipes (3m)", "120", "5,640.00", "Le Havre, FR"],
        ["SM-1005", "Electronic Control Units", "300", "420.00", "Antwerp, BE"],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            table.cell(row_idx, col_idx).text = val

    # Footer note
    doc.add_paragraph("")
    doc.add_paragraph("Authorized by: Captain R. Nakamura | Chief Mate: L. Petrov")

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
