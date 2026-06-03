"""
Reward Script: Insert two new rows after row 3 in shipping manifest table
Task ID: writer_tm_028
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Table row count changed from 6 to 8
  Component 2 (0.35): Rows 3-4 (0-indexed) are empty (the two inserted rows)
  Component 3 (0.35): Original data rows shifted correctly to rows 5-7 AND header/early rows preserved
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_028'

# Expected data from the original table
ORIGINAL_HEADER = ['Item No', 'Description', 'Quantity', 'Weight (kg)', 'Destination']
ORIGINAL_DATA_ROWS = [
    ['SM-1001', 'Industrial Bearings (SKF 6205)', '240', '1,080.00', 'Rotterdam, NL'],
    ['SM-1002', 'Copper Wire Spools (2.5mm)', '85', '3,400.00', 'Hamburg, DE'],
    ['SM-1003', 'Automotive Brake Pads (Set)', '500', '750.00', 'Felixstowe, UK'],
    ['SM-1004', 'Stainless Steel Pipes (3m)', '120', '5,640.00', 'Le Havre, FR'],
    ['SM-1005', 'Electronic Control Units', '300', '420.00', 'Antwerp, BE'],
]


def get_row_texts(table, row_idx):
    """Get list of stripped cell texts for a given row."""
    return [cell.text.strip() for cell in table.rows[row_idx].cells]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # Component 1: Table has 8 rows (was 6) — the core structural change (0.30 points)
    # This FAILS on initial (6 rows) and PASSES on golden (8 rows)
    try:
        if num_rows == 8 and num_cols == 5:
            print(f"PASS: Component 1 — Table is 8x5 (was 6x5) (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 1 — Expected 8 rows x 5 cols, found {num_rows} rows x {num_cols} cols")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Rows 3 and 4 (0-indexed) are empty inserted rows (0.35 points)
    # This FAILS on initial (rows 3-4 have SM-1003/SM-1004 data) and PASSES on golden (empty)
    try:
        if num_rows >= 5:
            row3_texts = get_row_texts(table, 3)
            row4_texts = get_row_texts(table, 4)
            row3_empty = all(t == '' for t in row3_texts)
            row4_empty = all(t == '' for t in row4_texts)
            if row3_empty and row4_empty:
                print(f"PASS: Component 2 — Rows 3 and 4 are empty (inserted rows) (0.35 pts)")
                total_score += 0.35
            elif row3_empty or row4_empty:
                print(f"FAIL: Component 2 — Only one of rows 3-4 is empty. Row 3: {row3_texts}, Row 4: {row4_texts}")
                total_score += 0.15  # partial: one row inserted
            else:
                print(f"FAIL: Component 2 — Neither row 3 nor row 4 is empty. Row 3: {row3_texts}, Row 4: {row4_texts}")
        else:
            print(f"FAIL: Component 2 — Not enough rows ({num_rows}) to check inserted rows")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Original data rows 3-5 (SM-1003, SM-1004, SM-1005) shifted to rows 5-7
    # AND header + first 2 data rows are still intact (compound check) (0.35 points)
    # This FAILS on initial (row 5 doesn't exist / data at wrong positions) and PASSES on golden
    try:
        if num_rows >= 8:
            # Check shifted rows (the task-introduced change)
            shifted_correct = 0
            for i, orig_idx in enumerate([2, 3, 4]):  # ORIGINAL_DATA_ROWS indices for SM-1003..SM-1005
                new_row_idx = 5 + i  # rows 5, 6, 7
                actual = get_row_texts(table, new_row_idx)
                expected = ORIGINAL_DATA_ROWS[orig_idx]
                if actual == expected:
                    shifted_correct += 1
                else:
                    print(f"  INFO: Row {new_row_idx} mismatch: expected {expected}, found {actual}")

            # Also verify header and first 2 data rows preserved (sub-condition, not standalone score)
            preserved_correct = 0
            if get_row_texts(table, 0) == ORIGINAL_HEADER:
                preserved_correct += 1
            if get_row_texts(table, 1) == ORIGINAL_DATA_ROWS[0]:
                preserved_correct += 1
            if get_row_texts(table, 2) == ORIGINAL_DATA_ROWS[1]:
                preserved_correct += 1

            total_checks = shifted_correct + preserved_correct  # out of 6
            if total_checks == 6:
                print(f"PASS: Component 3 — All shifted rows correct + preserved rows intact (0.35 pts)")
                total_score += 0.35
            elif shifted_correct >= 2:
                partial = round(0.35 * total_checks / 6, 2)
                print(f"FAIL: Component 3 — {shifted_correct}/3 shifted, {preserved_correct}/3 preserved ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — Shifted rows incorrect ({shifted_correct}/3)")
        else:
            print(f"FAIL: Component 3 — Not enough rows ({num_rows}) to verify shifted data")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
