"""
Reward Script: Set up a page break before 'Conclusion' paragraph and apply landscape page style
Task ID: writer_page_057
Domain: libreoffice_writer
Scoring:
  Component 1: Document has 2+ sections (section break added before Conclusion) — 0.3 pts
  Component 2: The section containing Conclusion has landscape orientation — 0.4 pts
  Component 3: Landscape section has correct A4 dimensions and margins — 0.3 pts
"""

import os
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'analysis_report'

# A4 landscape dimensions in EMU (1 inch = 914400 EMU, 1 cm = 360000 EMU)
# A4: 21.0cm x 29.7cm -> landscape: 29.7cm x 21.0cm
A4_LANDSCAPE_WIDTH_CM = 29.70
A4_LANDSCAPE_HEIGHT_CM = 21.00
MARGIN_CM = 2.54
TOLERANCE_CM = 0.15  # 1.5mm tolerance for rounding


def cm_to_emu(cm):
    return cm * 360000


def emu_to_cm(emu):
    return emu / 360000


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Insert a section break before the 'Conclusion' heading and set
    the conclusion section to landscape orientation (A4, same margins).
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must contain a 'Conclusion' heading
    conclusion_para_index = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith('Conclusion') and 'Heading' in para.style.name:
            conclusion_para_index = i
            break

    if conclusion_para_index is None:
        # Try matching just the text in any paragraph
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == 'Conclusion':
                conclusion_para_index = i
                break

    if conclusion_para_index is None:
        print("CRITICAL: 'Conclusion' paragraph not found — document structure is unexpected")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)

    # Component 1: Document has 2 sections (a new section break was inserted before Conclusion)
    # Initial doc has 1 section; golden doc should have 2 sections.
    # This FAILS on initial (1 section) and PASSES on golden (2 sections).
    try:
        if num_sections >= 2:
            # Also verify that there's a section break in a paragraph before or at the Conclusion paragraph
            # Find the index of the first paragraph containing a sectPr in its pPr (section break marker)
            section_break_para_index = -1
            for i, para in enumerate(doc.paragraphs):
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    section_break_para_index = i
                    break

            if section_break_para_index >= 0 and section_break_para_index <= conclusion_para_index:
                print(f"PASS: Component 1 — Document has {num_sections} sections with section break at paragraph [{section_break_para_index}], before Conclusion paragraph [{conclusion_para_index}] (0.3 pts)")
                total_score += 0.3
            elif section_break_para_index > conclusion_para_index:
                print(f"FAIL: Component 1 — Document has {num_sections} sections but section break is at paragraph [{section_break_para_index}], AFTER Conclusion paragraph [{conclusion_para_index}]")
            else:
                print(f"FAIL: Component 1 — Document has {num_sections} sections but no explicit section break paragraph found before Conclusion")
        else:
            print(f"FAIL: Component 1 — Expected 2+ sections, found {num_sections}; no section break before Conclusion")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: The last section (containing Conclusion and beyond) has landscape orientation
    # Initial doc has only portrait sections; golden doc's last section must be landscape.
    # This FAILS on initial (only 1 portrait section) and PASSES on golden (2nd section is landscape).
    try:
        if num_sections >= 2:
            # The section after the break (last section) should be landscape
            last_section = doc.sections[-1]
            if last_section.orientation == WD_ORIENT.LANDSCAPE:
                print(f"PASS: Component 2 — Last section (conclusion section) has landscape orientation (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 2 — Last section orientation is {last_section.orientation}, expected LANDSCAPE")
        else:
            print(f"FAIL: Component 2 — Only {num_sections} section(s) found; cannot check landscape orientation of conclusion section")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Landscape section has correct A4 dimensions and 2.54cm margins
    # Initial doc only has portrait section; golden last section should have correct A4 landscape dims + margins.
    # This FAILS on initial and PASSES on golden.
    try:
        if num_sections >= 2:
            last_section = doc.sections[-1]
            width_cm = emu_to_cm(last_section.page_width)
            height_cm = emu_to_cm(last_section.page_height)
            top_cm = emu_to_cm(last_section.top_margin)
            bottom_cm = emu_to_cm(last_section.bottom_margin)
            left_cm = emu_to_cm(last_section.left_margin)
            right_cm = emu_to_cm(last_section.right_margin)

            # Check A4 landscape dimensions: width ~29.7cm, height ~21.0cm
            width_ok = abs(width_cm - A4_LANDSCAPE_WIDTH_CM) <= TOLERANCE_CM
            height_ok = abs(height_cm - A4_LANDSCAPE_HEIGHT_CM) <= TOLERANCE_CM

            # Check margins: all 4 should be 2.54cm
            top_ok = abs(top_cm - MARGIN_CM) <= TOLERANCE_CM
            bottom_ok = abs(bottom_cm - MARGIN_CM) <= TOLERANCE_CM
            left_ok = abs(left_cm - MARGIN_CM) <= TOLERANCE_CM
            right_ok = abs(right_cm - MARGIN_CM) <= TOLERANCE_CM

            dims_ok = width_ok and height_ok
            margins_ok = top_ok and bottom_ok and left_ok and right_ok

            if dims_ok and margins_ok:
                print(f"PASS: Component 3 — Landscape section has A4 dimensions ({width_cm:.2f}cm x {height_cm:.2f}cm) and correct margins ({top_cm:.2f}cm) (0.3 pts)")
                total_score += 0.3
            elif dims_ok:
                print(f"FAIL: Component 3 — A4 dimensions OK but margins incorrect: top={top_cm:.2f}, bottom={bottom_cm:.2f}, left={left_cm:.2f}, right={right_cm:.2f} (expected 2.54cm each)")
            elif margins_ok:
                print(f"FAIL: Component 3 — Margins OK but A4 dimensions incorrect: {width_cm:.2f}cm x {height_cm:.2f}cm (expected {A4_LANDSCAPE_WIDTH_CM}cm x {A4_LANDSCAPE_HEIGHT_CM}cm)")
            else:
                print(f"FAIL: Component 3 — Both dimensions ({width_cm:.2f}cm x {height_cm:.2f}cm) and margins incorrect")
        else:
            print(f"FAIL: Component 3 — Only {num_sections} section(s) found; cannot check landscape section properties")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path on the VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
