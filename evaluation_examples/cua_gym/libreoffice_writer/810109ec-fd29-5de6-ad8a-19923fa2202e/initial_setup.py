"""
Initial Setup: Data analysis report - before page break and landscape orientation for Conclusion
Task ID: writer_page_057
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION_START

WORKDIR = '/home/user'
TASK_ID = 'writer_page_057'
OUTPUT = f'{WORKDIR}/Desktop/analysis_report.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # ---- Page setup: A4, portrait, all margins 2.54cm ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ---- Title ----
    title = doc.add_heading('Quarterly Business Performance Analysis Report', level=0)
    title.paragraph_format.space_after = Pt(12)

    # ---- Section 1: Executive Summary ----
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report presents a comprehensive analysis of the company\'s quarterly '
        'business performance metrics for Q3 2024. Key performance indicators have '
        'been evaluated across revenue, operations, customer satisfaction, and market '
        'positioning. Overall, the company demonstrated a 12.4% revenue growth compared '
        'to the same quarter last year, driven primarily by expanded product offerings '
        'and increased market penetration in the Asia-Pacific region.'
    )
    doc.add_paragraph(
        'Operational efficiency metrics improved by 8.7%, while customer retention '
        'rates remained stable at 94.2%. The report identifies three core areas '
        'for strategic focus in the upcoming quarter: talent acquisition, digital '
        'transformation, and supply chain optimization.'
    )

    # ---- Section 2: Introduction ----
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'The purpose of this analysis is to provide stakeholders with a detailed '
        'overview of business performance for Q3 2024 (July – September). Data was '
        'collected from internal systems including ERP, CRM, and financial reporting '
        'platforms. The methodology employed adheres to the company\'s standard '
        'performance benchmarking framework established in 2022.'
    )
    doc.add_paragraph(
        'Data sources include: monthly revenue reports from the Finance Department, '
        'operational KPIs from the Operations Excellence team, customer survey data '
        'compiled by the Customer Experience division, and market intelligence reports '
        'from the Strategy & Business Development group.'
    )
    doc.add_paragraph(
        'Unless otherwise noted, all financial figures are expressed in USD and '
        'represent consolidated results across all subsidiaries and operating regions. '
        'Year-over-year (YoY) comparisons are based on Q3 2023 actuals.'
    )

    # ---- Section 3: Data Analysis ----
    doc.add_heading('Data Analysis', level=1)
    doc.add_heading('3.1 Revenue Performance', level=2)
    doc.add_paragraph(
        'Total revenue for Q3 2024 reached $48.7 million, compared to $43.3 million '
        'in Q3 2023, representing a year-over-year increase of 12.4%. Product revenue '
        'accounted for 68% of total revenue ($33.1M), while service revenue contributed '
        '32% ($15.6M). The North America region remained the largest revenue contributor '
        'at 52% ($25.3M), followed by EMEA at 28% ($13.6M) and APAC at 20% ($9.8M).'
    )

    # Revenue table
    table1 = doc.add_table(rows=5, cols=4)
    table1.style = 'Table Grid'
    headers1 = ['Region', 'Q3 2023 ($M)', 'Q3 2024 ($M)', 'YoY Growth (%)']
    for col_idx, h in enumerate(headers1):
        cell = table1.cell(0, col_idx)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data_rows1 = [
        ['North America', '22.1', '25.3', '+14.5%'],
        ['EMEA', '12.4', '13.6', '+9.7%'],
        ['APAC', '7.5', '9.8', '+30.7%'],
        ['Total', '43.3', '48.7', '+12.4%'],
    ]
    for row_idx, row_data in enumerate(data_rows1, 1):
        for col_idx, val in enumerate(row_data):
            table1.cell(row_idx, col_idx).text = val
    doc.add_paragraph('')

    doc.add_heading('3.2 Operational Efficiency', level=2)
    doc.add_paragraph(
        'Operational costs as a percentage of revenue declined from 61.2% in Q3 2023 '
        'to 56.5% in Q3 2024, reflecting efficiency gains from automation initiatives '
        'and workforce optimization programs. The average order fulfillment time '
        'decreased from 4.8 days to 3.9 days, a 18.8% improvement. Inventory turnover '
        'ratio improved from 5.2x to 6.1x annually.'
    )
    doc.add_paragraph(
        'The Six Sigma continuous improvement program implemented in Q1 2024 contributed '
        'to a 23% reduction in process defects. Cross-functional collaboration index '
        'scores rose from 72/100 to 81/100, indicating stronger inter-departmental '
        'alignment and knowledge sharing.'
    )

    doc.add_heading('3.3 Customer Metrics', level=2)
    doc.add_paragraph(
        'Net Promoter Score (NPS) for Q3 2024 reached 67, up from 59 in Q3 2023. '
        'Customer acquisition cost (CAC) decreased by 11% from $342 to $304, while '
        'customer lifetime value (CLV) increased by 18% from $4,850 to $5,723. '
        'Churn rate remained stable at 5.8% annually, with the highest retention '
        'rates observed in the Enterprise segment (97.1%) compared to SMB (91.4%).'
    )

    # ---- Section 4: Findings ----
    doc.add_heading('Findings', level=1)
    doc.add_paragraph(
        'Based on the data analysis conducted, the following key findings have been '
        'identified for Q3 2024:'
    )
    doc.add_paragraph(
        'Revenue Growth Drivers: The 12.4% YoY revenue growth was primarily driven '
        'by APAC expansion (30.7% growth), successful launch of three new product lines '
        'in July 2024, and a 19% increase in average contract value for enterprise customers.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Operational Improvements: Process automation investments totaling $2.1M in '
        'H1 2024 yielded cost savings of approximately $3.7M in Q3 alone, delivering '
        'a 176% return on investment within the same fiscal year.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Customer Experience: The NPS improvement from 59 to 67 reflects successful '
        'implementation of the Customer Success Program launched in Q2 2024. Response '
        'time improvements (avg. from 4.2h to 1.8h) are the primary driver of this gain.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Risk Factors: Foreign exchange volatility impacted APAC revenues by an '
        'estimated -2.3%, and global supply chain disruptions resulted in delayed '
        'deliveries for approximately 340 orders (representing 1.2% of total order volume).',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Market Position: The company\'s market share in the North American mid-market '
        'segment grew from 8.4% to 9.7%, while European market share remained flat '
        'at 5.2% despite favorable macroeconomic conditions in Germany and the Netherlands.',
        style='List Bullet'
    )

    doc.add_heading('4.1 Comparative Performance', level=2)
    doc.add_paragraph(
        'Compared to industry benchmarks, the company outperformed sector averages '
        'across 7 of 10 key performance indicators. Particularly strong performance '
        'was noted in revenue per employee ($187K vs. sector average $142K), gross '
        'margin (44.7% vs. sector average 38.9%), and NPS (67 vs. sector average 52).'
    )

    # Comparison table
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    headers2 = ['KPI', 'Company Q3 2024', 'Industry Average']
    for col_idx, h in enumerate(headers2):
        cell = table2.cell(0, col_idx)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data_rows2 = [
        ['Revenue per Employee ($K)', '187', '142'],
        ['Gross Margin (%)', '44.7%', '38.9%'],
        ['NPS Score', '67', '52'],
        ['Customer Churn (%)', '5.8%', '8.2%'],
    ]
    for row_idx, row_data in enumerate(data_rows2, 1):
        for col_idx, val in enumerate(row_data):
            table2.cell(row_idx, col_idx).text = val
    doc.add_paragraph('')

    # ---- Section 5: Conclusion ----
    # NO page break before Conclusion, NO landscape orientation (initial state)
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The Q3 2024 business performance analysis confirms the company is on a strong '
        'growth trajectory, driven by strategic investments in technology, talent, and '
        'market expansion. Revenue growth of 12.4% YoY demonstrates the effectiveness '
        'of the current go-to-market strategy, while operational efficiency gains '
        'validate the process improvement programs initiated in early 2024.'
    )
    doc.add_paragraph(
        'The significant APAC growth (30.7%) signals a maturing regional operation with '
        'substantial future potential. Management recommends continued investment in APAC '
        'infrastructure and talent to sustain and accelerate this momentum. Simultaneously, '
        'the stagnant European market share requires strategic review and targeted '
        'intervention in Q4 2024.'
    )
    doc.add_paragraph(
        'Looking ahead to Q4 2024, three strategic priorities are recommended: '
        '(1) Accelerate digital transformation initiatives to capture an additional '
        '2-3% operational cost reduction; (2) Execute the planned talent acquisition '
        'drive to support APAC expansion with an estimated 45 new headcount; and '
        '(3) Strengthen European partnerships and channel strategy to regain competitive '
        'market positioning.'
    )
    doc.add_paragraph(
        'The company is well-positioned to achieve full-year 2024 revenue targets of '
        '$195M and operating margin targets of 18.5%, assuming no significant macro '
        'disruptions in Q4. Risk mitigation strategies for currency volatility and '
        'supply chain exposure have been reviewed and are deemed adequate for current '
        'exposure levels.'
    )

    # ---- Section 6: References ----
    doc.add_heading('References', level=1)
    doc.add_paragraph(
        'Internal Finance Department. (2024). Quarterly Revenue Report Q3 2024. '
        'Company Internal Publication.',
        style='List Number'
    )
    doc.add_paragraph(
        'Operations Excellence Team. (2024). Q3 2024 Operational KPI Dashboard. '
        'Company Internal Publication.',
        style='List Number'
    )
    doc.add_paragraph(
        'Customer Experience Division. (2024). Customer Satisfaction Survey Results '
        'Q3 2024. Company Internal Publication.',
        style='List Number'
    )
    doc.add_paragraph(
        'Strategy & Business Development Group. (2024). Market Intelligence Report '
        'Asia-Pacific Region Q3 2024. Company Internal Publication.',
        style='List Number'
    )
    doc.add_paragraph(
        'Industry Research Institute. (2024). Global Technology Sector Benchmarking '
        'Report 2024. Industry Research Institute Publications.',
        style='List Number'
    )
    doc.add_paragraph(
        'Supply Chain Risk Management Committee. (2024). Q3 2024 Supply Chain '
        'Disruption Impact Assessment. Company Internal Publication.',
        style='List Number'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
