"""
Initial Setup: Workflow Overview document with heading but no diagram shapes.
Task ID: writer_obj_054
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'workflow_doc'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set up page: A4 portrait
    section = doc.sections[0]
    section.page_width = Inches(8.27)    # A4 width
    section.page_height = Inches(11.69)  # A4 height
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Add heading 'Workflow Overview'
    heading = doc.add_heading('Workflow Overview', level=1)
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add some placeholder text below heading (realistic content, no diagram)
    para = doc.add_paragraph(
        'This document outlines the standard workflow process for data handling within the organization. '
        'The process involves receiving input data, processing it through internal validation pipelines, '
        'and delivering the resulting output to downstream systems.'
    )
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    para2 = doc.add_paragraph(
        'The workflow ensures data integrity at each stage, with proper error handling and logging mechanisms '
        'to track the status of each transaction from input receipt to final output delivery.'
    )
    para2.paragraph_format.space_before = Pt(6)
    para2.paragraph_format.space_after = Pt(6)

    # Save the document (no shapes/diagram)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
