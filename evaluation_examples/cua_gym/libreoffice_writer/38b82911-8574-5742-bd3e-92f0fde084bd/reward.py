"""
Reward Script: Convert plain-text attendance record into formatted table with conditional formatting
Task ID: writer_hr_046
Domain: libreoffice_writer
Scoring:
  - Component 1: Table exists with correct dimensions (16 rows x 5 cols) — 0.3 pts
  - Component 2: Data content matches expected employee records — 0.25 pts
  - Component 3: Header row is formatted distinctly (bold) — 0.15 pts
  - Component 4: Attendance rates below 90% have red bold text — 0.3 pts
"""

import os
from docx import Document
from docx.shared import RGBColor
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_046'

# Expected employee data (name, department, days_present, days_absent, attendance_rate)
EXPECTED_HEADERS = ['Employee Name', 'Department', 'Days Present', 'Days Absent', 'Attendance Rate']

EXPECTED_EMPLOYEES = [
    ('Sarah Chen', 'Engineering', '228', '12', '95.0%'),
    ('Marcus Johnson', 'Marketing', '210', '30', '87.5%'),
    ('Priya Patel', 'Finance', '232', '8', '96.7%'),
    ('David Kim', 'Engineering', '205', '35', '85.4%'),
    ('Elena Rodriguez', 'Human Resources', '235', '5', '97.9%'),
    ("James O'Brien", 'Sales', '220', '20', '91.7%'),
    ('Aisha Mohammed', 'Finance', '238', '2', '99.2%'),
    ('Robert Taylor', 'Marketing', '198', '42', '82.5%'),
    ('Lin Wei', 'Engineering', '230', '10', '95.8%'),
    ('Sophie Martin', 'Sales', '212', '28', '88.3%'),
    ('Carlos Gutierrez', 'Operations', '225', '15', '93.8%'),
    ('Fatima Al-Rashid', 'Human Resources', '215', '25', '89.6%'),
    ('Thomas Anderson', 'Operations', '233', '7', '97.1%'),
    ('Yuki Tanaka', 'Finance', '208', '32', '86.7%'),
    ('Michael Foster', 'Sales', '222', '18', '92.5%'),
]

# Employees with attendance rate < 90% (should be red+bold)
BELOW_90_INDICES = [1, 3, 7, 9, 11, 13]  # 0-based row indices in employee list


def color_distance(c1, c2):
    """Euclidean distance between two RGB colors."""
    return sqrt(sum((a - b) ** 2 for a, b in zip(c1, c2)))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with correct dimensions (0.3 points)
    # Initial file has NO tables — only golden file should have one
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 16 and num_cols == 5:
                print(f"PASS: Component 1 — Table exists with 16 rows x 5 cols (0.3 pts)")
                total_score += 0.3
            elif num_rows >= 2 and num_cols == 5:
                # Partial: table exists with right columns but wrong row count
                partial = 0.15
                print(f"PARTIAL: Component 1 — Table has {num_rows} rows x {num_cols} cols, expected 16x5 ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 1 — Table has {num_rows} rows x {num_cols} cols, expected 16x5")
        else:
            print(f"FAIL: Component 1 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Gate: need at least one table to continue
    if len(doc.tables) == 0:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    table = doc.tables[0]

    # Component 2: Data content matches expected employee records (0.25 points)
    # Initial file has data in paragraphs, not in table cells — so this checks the table content
    try:
        # Check headers
        header_cells = [table.cell(0, c).text.strip() for c in range(min(5, len(table.columns)))]
        headers_match = all(
            h.lower() == e.lower()
            for h, e in zip(header_cells, EXPECTED_HEADERS)
        ) if len(header_cells) == 5 else False

        # Check employee data
        correct_employees = 0
        total_employees = min(len(table.rows) - 1, 15)
        for i in range(total_employees):
            row_idx = i + 1  # skip header
            try:
                row_data = [table.cell(row_idx, c).text.strip() for c in range(5)]
                expected = EXPECTED_EMPLOYEES[i]
                if all(r == e for r, e in zip(row_data, expected)):
                    correct_employees += 1
            except Exception:
                pass

        if headers_match and correct_employees >= 14:
            print(f"PASS: Component 2 — Headers correct, {correct_employees}/15 employees match (0.25 pts)")
            total_score += 0.25
        elif headers_match or correct_employees >= 10:
            partial = 0.125
            print(f"PARTIAL: Component 2 — Headers={'match' if headers_match else 'mismatch'}, {correct_employees}/15 employees match ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Headers={'match' if headers_match else 'mismatch'}, {correct_employees}/15 employees match")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header row formatted distinctly — bold (0.15 points)
    # Initial file has NO table, so header formatting only exists in golden
    try:
        header_row = table.rows[0]
        bold_count = 0
        for ci in range(min(5, len(table.columns))):
            cell = header_row.cells[ci]
            cell_has_bold = False
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.bold:
                        cell_has_bold = True
                        break
                if cell_has_bold:
                    break
            if cell_has_bold:
                bold_count += 1

        if bold_count >= 4:
            print(f"PASS: Component 3 — Header row has {bold_count}/5 bold cells (0.15 pts)")
            total_score += 0.15
        elif bold_count >= 2:
            partial = 0.075
            print(f"PARTIAL: Component 3 — Header row has {bold_count}/5 bold cells ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Header row has {bold_count}/5 bold cells")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Attendance rates below 90% have red bold text (0.3 points)
    # There are 6 employees with rate < 90%. Each correctly formatted = 0.05 pts
    try:
        red_bold_correct = 0
        red_bold_false_positive = 0
        total_check = min(len(table.rows) - 1, 15)

        for i in range(total_check):
            row_idx = i + 1  # skip header
            should_be_red_bold = i in BELOW_90_INDICES
            try:
                rate_cell = table.cell(row_idx, 4)
                is_red = False
                is_bold = False
                for para in rate_cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            if run.font.color.rgb is not None:
                                # Check if color is close to red (FF0000)
                                rgb = run.font.color.rgb
                                dist = color_distance(
                                    (rgb[0], rgb[1], rgb[2]) if hasattr(rgb, '__getitem__') else (int(str(rgb)[:2], 16), int(str(rgb)[2:4], 16), int(str(rgb)[4:6], 16)),
                                    (255, 0, 0)
                                )
                                if dist < 80:
                                    is_red = True
                            if run.font.bold:
                                is_bold = True

                if should_be_red_bold:
                    if is_red and is_bold:
                        red_bold_correct += 1
                else:
                    if is_red or is_bold:
                        red_bold_false_positive += 1
            except Exception:
                pass

        # Score: each correctly red+bold cell is worth 0.05, minus false positives
        points = red_bold_correct * 0.05
        if red_bold_false_positive > 0:
            points = max(0, points - red_bold_false_positive * 0.025)
        points = min(points, 0.3)

        if red_bold_correct >= 5 and red_bold_false_positive == 0:
            print(f"PASS: Component 4 — {red_bold_correct}/6 below-90% cells are red+bold, {red_bold_false_positive} false positives (0.3 pts)")
            total_score += 0.3
        elif red_bold_correct >= 1:
            print(f"PARTIAL: Component 4 — {red_bold_correct}/6 below-90% cells are red+bold, {red_bold_false_positive} false positives ({points} pts)")
            total_score += points
        else:
            print(f"FAIL: Component 4 — {red_bold_correct}/6 below-90% cells are red+bold")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
