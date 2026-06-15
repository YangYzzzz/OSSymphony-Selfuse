"""
Reward Script: Create a grading rubric table in LibreOffice Writer
Task ID: writer_tbl_044
Domain: libreoffice_writer
Scoring:
  - Component 1: Table exists with 4 rows and 5 columns (0.3 pts)
  - Component 2: Header row has correct column names (0.3 pts)
  - Component 3: Header row cells have green background (0.2 pts)
  - Component 4: Header row text is bold (0.1 pts)
  - Component 5: First column rows 2-4 have correct entries (0.1 pts)
"""

import os
from docx import Document
from docx.oxml.ns import qn
from math import sqrt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_044'
FILE_PATH = f'{WORKDIR}/grading_rubric.docx'

EXPECTED_HEADERS = ['Criteria', 'Excellent (4)', 'Good (3)', 'Satisfactory (2)', 'Needs Work (1)']
EXPECTED_FIRST_COL = ['Content', 'Organization', 'Grammar']

# Green background used in golden: 00B050 — accept any reasonably green fill
# We check: green channel clearly dominant and not white/transparent
def is_green_color(fill_hex):
    """Check if a fill hex color is a shade of green."""
    if not fill_hex or fill_hex.upper() in ('AUTO', 'FFFFFF', '000000', 'NONE', ''):
        return False
    try:
        r = int(fill_hex[0:2], 16)
        g = int(fill_hex[2:4], 16)
        b = int(fill_hex[4:6], 16)
        # Green dominant: g > r and g > b and g must be reasonably bright
        # Also accept if it's close to known green 00B050 (R=0, G=176, B=80)
        known_green = (0, 176, 80)
        dist = sqrt((r - known_green[0])**2 + (g - known_green[1])**2 + (b - known_green[2])**2)
        if dist < 80:
            return True
        # General green: green channel dominant
        if g > r + 30 and g > b + 30 and g > 80:
            return True
        return False
    except Exception:
        return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: file title still present (non-scoring gate)
    title_present = any(p.text.strip() == 'Essay Grading Rubric' for p in doc.paragraphs)
    if not title_present:
        print("WARN: Title paragraph 'Essay Grading Rubric' not found — may indicate wrong file")

    # Component 1: Table exists with 4 rows and 5 columns (0.3 points)
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 1 — No table found in document")
        else:
            table = doc.tables[0]
            row_count = len(table.rows)
            col_count = len(table.columns)
            if row_count == 4 and col_count == 5:
                print(f"PASS: Component 1 — Table has correct dimensions {row_count}x{col_count} (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 1 — Expected 4x5 table, found {row_count}x{col_count}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Exit early if no table (remaining checks depend on table)
    if len(doc.tables) == 0:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    table = doc.tables[0]

    # Component 2: Header row has correct column names (0.3 points)
    try:
        header_row = table.rows[0]
        header_texts = [cell.text.strip() for cell in header_row.cells]
        if header_texts == EXPECTED_HEADERS:
            print(f"PASS: Component 2 — Header row matches expected columns {header_texts} (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Expected headers {EXPECTED_HEADERS}, found {header_texts}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header row cells have green background (0.2 points)
    try:
        header_row = table.rows[0]
        green_count = 0
        fill_values = []
        for cell in header_row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            cell_fill = None
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    cell_fill = shd.get(qn('w:fill'))
            fill_values.append(cell_fill)
            if cell_fill and is_green_color(cell_fill):
                green_count += 1
        if green_count == 5:
            print(f"PASS: Component 3 — All 5 header cells have green background (fills: {fill_values}) (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — Expected 5 green header cells, found {green_count} (fills: {fill_values})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Header row text is bold (0.1 points)
    try:
        header_row = table.rows[0]
        bold_count = 0
        for cell in header_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip() and run.font.bold is True:
                        bold_count += 1
        if bold_count >= 5:
            print(f"PASS: Component 4 — Header row text is bold ({bold_count} bold runs found) (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 4 — Expected all 5 header cells bold, found {bold_count} bold runs")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: First column rows 2-4 have correct entries (0.1 points)
    try:
        first_col_entries = []
        for r_i in range(1, 4):
            if r_i < len(table.rows):
                first_col_entries.append(table.rows[r_i].cells[0].text.strip())
        if first_col_entries == EXPECTED_FIRST_COL:
            print(f"PASS: Component 5 — First column rows 2-4: {first_col_entries} (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 5 — Expected {EXPECTED_FIRST_COL}, found {first_col_entries}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
