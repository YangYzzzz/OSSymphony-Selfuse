"""
Initial Setup: Plain document with tip text as unformatted paragraph
Task ID: writer_tech_039
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_039'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    heading = doc.add_heading('Deployment Best Practices', level=1)

    # Introductory paragraph
    p1 = doc.add_paragraph(
        'Modern software deployment requires careful management of configurations, '
        'secrets, and environment-specific settings. This guide covers key strategies '
        'for maintaining reliable deployment pipelines across development, staging, '
        'and production environments.'
    )

    # Section heading
    doc.add_heading('Configuration Management', level=2)

    p2 = doc.add_paragraph(
        'One of the most critical aspects of deployment is ensuring that configuration '
        'values are correctly set for each target environment. Hardcoding values directly '
        'into source code creates maintenance challenges and increases the risk of '
        'accidentally deploying development settings to production.'
    )

    # The tip paragraph - plain, unformatted text (this is what the agent must format)
    tip_para = doc.add_paragraph(
        'TIP: Use environment variables to manage different deployment configurations.'
    )

    # More content after the tip
    doc.add_heading('Continuous Integration', level=2)

    p3 = doc.add_paragraph(
        'Automated build and test pipelines should validate configuration consistency '
        'before any deployment proceeds. Tools like Jenkins, GitHub Actions, and GitLab CI '
        'can be configured to run environment-specific test suites that verify correct '
        'variable bindings and service connectivity.'
    )

    p4 = doc.add_paragraph(
        'When setting up CI/CD pipelines, ensure that sensitive credentials are stored '
        'in the pipeline secret manager rather than in repository files. Rotate secrets '
        'on a regular schedule and audit access logs for any unauthorized usage patterns.'
    )

    doc.add_heading('Rollback Strategies', level=2)

    p5 = doc.add_paragraph(
        'Every deployment should have a documented rollback plan. Blue-green deployments '
        'and canary releases provide mechanisms for gradual rollout and quick reversion '
        'if issues are detected in production. Monitor error rates and latency metrics '
        'closely during the first 30 minutes after each release.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
