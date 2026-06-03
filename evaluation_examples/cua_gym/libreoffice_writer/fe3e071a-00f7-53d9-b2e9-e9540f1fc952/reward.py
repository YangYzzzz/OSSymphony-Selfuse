"""
Reward Script: Property Inspection Report in LibreOffice Writer
Task ID: writer_wf_071
Domain: libreoffice_writer
Scoring:
  C1 (0.10) - Title 'Property Inspection Report' with Title style
  C2 (0.15) - Property details table with 5 required fields
  C3 (0.25) - Six Heading 2 inspection sections
  C4 (0.25) - Each inspection section has a 3-row condition table (Item/Condition/Notes)
  C5 (0.15) - Summary section with overall rating and recommendations
  C6 (0.10) - Inspector signature block at end
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_071'

REQUIRED_SECTIONS = ['Exterior', 'Roof', 'Plumbing', 'Electrical', 'HVAC', 'Interior']
VALID_CONDITIONS = {'good', 'fair', 'poor'}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect paragraph data for reuse
    paragraphs = doc.paragraphs
    tables = doc.tables

    # Component 1: Title 'Property Inspection Report' with Title style (0.10 points)
    try:
        title_found = False
        for p in paragraphs:
            if p.style and p.style.name == 'Title' and 'property inspection report' in p.text.lower():
                title_found = True
                break
        if title_found:
            print(f"PASS: Component 1 — Title found with Title style (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — Title 'Property Inspection Report' with Title style not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Property details table with 5 required fields (0.15 points)
    # The task specifies: Address, Owner, Inspector, Date, Report Number
    try:
        required_fields = {'address', 'owner', 'inspector', 'date', 'report number'}
        prop_table_found = False
        fields_found = set()

        for t in tables:
            # Look for a table with 2 columns that has the property detail fields
            if len(t.columns) == 2:
                row_labels = set()
                for row in t.rows:
                    label = row.cells[0].text.strip().lower()
                    row_labels.add(label)
                matched = required_fields & row_labels
                if len(matched) >= 4:  # At least 4 of 5 fields found
                    fields_found = matched
                    prop_table_found = True
                    break

        if prop_table_found and len(fields_found) >= 5:
            print(f"PASS: Component 2 — Property details table with all 5 fields (0.15 pts)")
            total_score += 0.15
        elif prop_table_found and len(fields_found) >= 4:
            print(f"PARTIAL: Component 2 — Property details table with {len(fields_found)}/5 fields (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — Property details table not found or missing fields. Found: {fields_found}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Six Heading 2 inspection sections (0.25 points)
    # Required: Exterior, Roof, Plumbing, Electrical, HVAC, Interior
    try:
        heading2_texts = []
        for p in paragraphs:
            if p.style and p.style.name == 'Heading 2':
                heading2_texts.append(p.text.strip().lower())

        sections_found = []
        for section_name in REQUIRED_SECTIONS:
            for h2 in heading2_texts:
                if section_name.lower() in h2:
                    sections_found.append(section_name)
                    break

        num_found = len(sections_found)
        if num_found == 6:
            print(f"PASS: Component 3 — All 6 inspection sections found as Heading 2 (0.25 pts)")
            total_score += 0.25
        elif num_found >= 4:
            partial = round(0.25 * (num_found / 6), 2)
            print(f"PARTIAL: Component 3 — {num_found}/6 inspection sections found ({partial} pts)")
            total_score += partial
        elif num_found >= 1:
            partial = round(0.25 * (num_found / 6), 2)
            print(f"PARTIAL: Component 3 — {num_found}/6 inspection sections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No inspection sections found as Heading 2. H2 headings: {heading2_texts}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Each inspection section has a condition table with 3 items (0.25 points)
    # Tables should have 3 columns (Item, Condition, Notes) and at least 3 data rows + 1 header row
    # Each data row's Condition column should contain Good/Fair/Poor
    try:
        # Find 3-column tables that look like condition tables
        condition_tables_count = 0
        for t in tables:
            if len(t.columns) == 3 and len(t.rows) >= 4:  # header + 3 data rows
                # Check header row
                headers = [cell.text.strip().lower() for cell in t.rows[0].cells]
                if 'item' in headers and 'condition' in headers and 'notes' in headers:
                    # Check that data rows have valid conditions
                    cond_col_idx = headers.index('condition')
                    valid_rows = 0
                    for row in t.rows[1:]:
                        cond_val = row.cells[cond_col_idx].text.strip().lower()
                        if cond_val in VALID_CONDITIONS:
                            valid_rows += 1
                    if valid_rows >= 3:
                        condition_tables_count += 1

        if condition_tables_count >= 6:
            print(f"PASS: Component 4 — All 6 condition tables found with valid items (0.25 pts)")
            total_score += 0.25
        elif condition_tables_count >= 1:
            partial = round(0.25 * (min(condition_tables_count, 6) / 6), 2)
            print(f"PARTIAL: Component 4 — {condition_tables_count}/6 condition tables found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — No valid condition tables (Item/Condition/Notes with 3 items) found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Summary section with overall rating and recommendations (0.15 points)
    try:
        # Check for a Summary heading
        summary_heading_found = False
        summary_heading_idx = -1
        for i, p in enumerate(paragraphs):
            if p.style and 'Heading' in (p.style.name or '') and 'summary' in p.text.strip().lower():
                summary_heading_found = True
                summary_heading_idx = i
                break

        # Check for overall rating text after the summary heading
        has_rating = False
        has_recommendations = False
        if summary_heading_found:
            for p in paragraphs[summary_heading_idx + 1:]:
                text_lower = p.text.lower()
                if 'overall rating' in text_lower or 'overall condition' in text_lower or 'rating' in text_lower:
                    has_rating = True
                if 'recommend' in text_lower:
                    has_recommendations = True

        if summary_heading_found and has_rating and has_recommendations:
            print(f"PASS: Component 5 — Summary with rating and recommendations (0.15 pts)")
            total_score += 0.15
        elif summary_heading_found and (has_rating or has_recommendations):
            print(f"PARTIAL: Component 5 — Summary found but missing {'recommendations' if not has_recommendations else 'rating'} (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 5 — Summary section not found or missing rating/recommendations")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Inspector signature block at end (0.10 points)
    # Check that near the end of the document there is a signature line and inspector info
    try:
        # Look at the last 10 paragraphs for signature indicators
        tail_paras = paragraphs[-10:] if len(paragraphs) > 10 else paragraphs
        tail_text = ' '.join(p.text for p in tail_paras).lower()

        has_signature_line = any(
            '___' in p.text or 'signature' in p.text.lower()
            for p in tail_paras
        )
        has_inspector_name = any(
            'inspector' in p.text.lower() or 'certified' in p.text.lower() or 'cpi' in p.text.lower()
            for p in tail_paras
        )

        if has_signature_line and has_inspector_name:
            print(f"PASS: Component 6 — Inspector signature block found (0.10 pts)")
            total_score += 0.10
        elif has_signature_line or has_inspector_name:
            print(f"PARTIAL: Component 6 — Partial signature block (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — No signature block found at end of document")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
