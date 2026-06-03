"""
Initial Setup: Writer document with References section, single-column layout
Task ID: writer_fs_066
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_066'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('Urban Transit Efficiency Study: Metropolitan Area Analysis', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    abstract = doc.add_paragraph(
        'This study examines the operational efficiency of public transit systems '
        'across twelve metropolitan areas in North America. Using ridership data '
        'collected between 2021 and 2024, we evaluate key performance indicators '
        'including on-time performance, cost per passenger mile, and fleet '
        'utilization rates. Our findings suggest that cities investing in dedicated '
        'bus rapid transit corridors achieved 23% higher ridership growth compared '
        'to conventional bus route expansions.'
    )

    # --- Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Public transportation remains a cornerstone of urban mobility strategy. '
        'As metropolitan populations continue to grow, transit agencies face '
        'mounting pressure to deliver efficient, reliable service while managing '
        'constrained budgets. The Federal Transit Administration reported that '
        'U.S. transit systems carried approximately 9.6 billion passenger trips '
        'in 2023, representing a significant recovery from pandemic-era lows.'
    )
    doc.add_paragraph(
        'This paper contributes to the growing body of literature on transit '
        'performance measurement by introducing a composite efficiency index (CEI) '
        'that integrates operational, financial, and ridership metrics into a '
        'single comparable score. We apply this index to twelve metropolitan '
        'transit agencies and identify common factors associated with higher '
        'efficiency scores.'
    )

    # --- Methodology ---
    doc.add_heading('2. Methodology', level=1)
    doc.add_paragraph(
        'Data were sourced from the National Transit Database (NTD) annual reports '
        'for fiscal years 2021 through 2024. The twelve agencies selected represent '
        'a mix of large (population > 2 million), medium (500,000 - 2 million), '
        'and small (< 500,000) metropolitan areas to ensure geographic and '
        'demographic diversity.'
    )
    doc.add_heading('2.1 Composite Efficiency Index', level=2)
    doc.add_paragraph(
        'The CEI is computed as a weighted average of three normalized sub-indices:'
    )
    doc.add_paragraph('On-Time Performance Index (OTPI) — weight: 0.35', style='List Bullet')
    doc.add_paragraph('Cost Efficiency Index (CECI) — weight: 0.40', style='List Bullet')
    doc.add_paragraph('Fleet Utilization Index (FUI) — weight: 0.25', style='List Bullet')
    doc.add_paragraph(
        'Each sub-index is scaled to a 0-100 range using min-max normalization '
        'across the sample. The weighting scheme was validated through a Delphi '
        'process involving fourteen transit planning professionals.'
    )

    # --- Results table ---
    doc.add_heading('3. Results', level=1)
    doc.add_paragraph(
        'Table 1 summarizes the CEI scores and sub-index values for the twelve '
        'agencies in the 2024 reporting year.'
    )

    table = doc.add_table(rows=13, cols=5)
    table.style = 'Table Grid'
    headers = ['Agency', 'OTPI', 'CECI', 'FUI', 'CEI']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Metro Transit Authority', '82.4', '76.1', '88.3', '81.2'],
        ['Bay Area Rapid Transit', '79.8', '71.5', '91.0', '79.1'],
        ['Capital Metro Services', '85.1', '80.3', '79.6', '81.9'],
        ['Lakeside Transit Corp.', '74.2', '83.7', '72.5', '78.0'],
        ['Pacific Coast Transport', '88.6', '68.9', '85.1', '79.4'],
        ['Heartland Bus Network', '71.3', '89.2', '66.8', '78.3'],
        ['Atlantic Commuter Rail', '90.1', '62.4', '93.7', '79.0'],
        ['Mountain Valley Transit', '68.5', '77.8', '70.2', '73.1'],
        ['Great Plains Express', '76.9', '85.6', '74.0', '80.2'],
        ['Delta Metro Authority', '83.7', '73.2', '82.5', '79.0'],
        ['Northern Light Rail', '91.2', '66.8', '89.4', '80.4'],
        ['Sunbelt Rapid Transit', '87.3', '79.1', '81.6', '82.5'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Discussion ---
    doc.add_heading('4. Discussion', level=1)
    doc.add_paragraph(
        'The results reveal that no single agency dominates across all three '
        'sub-indices. Sunbelt Rapid Transit achieved the highest overall CEI '
        'score (82.5), driven by balanced performance across operational and '
        'financial dimensions. Agencies with dedicated right-of-way infrastructure '
        'consistently scored higher on OTPI, while those with newer fleets '
        'showed stronger FUI performance.'
    )
    doc.add_paragraph(
        'A notable finding is the inverse relationship between OTPI and CECI '
        'observed in several agencies. Atlantic Commuter Rail, for example, '
        'achieved the highest OTPI (90.1) but the lowest CECI (62.4), suggesting '
        'that maintaining high on-time performance may come at a significant '
        'cost premium for rail-based systems.'
    )

    # --- Conclusion ---
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        'The Composite Efficiency Index provides a useful framework for '
        'benchmarking transit agency performance. Our analysis demonstrates '
        'that balanced investment across operational reliability, cost management, '
        'and fleet modernization yields the strongest overall efficiency outcomes. '
        'Future work should extend this framework to include passenger satisfaction '
        'metrics and environmental impact indicators.'
    )

    # --- References (the final section before appendix would be added) ---
    doc.add_heading('References', level=1)
    doc.add_paragraph(
        'American Public Transportation Association. (2024). '
        'Public Transportation Fact Book (75th ed.). Washington, DC: APTA.'
    )
    doc.add_paragraph(
        'Chen, W., & Rodriguez, M. (2023). Benchmarking urban transit '
        'performance: A multi-criteria approach. Transportation Research '
        'Part A: Policy and Practice, 168, 45-62.'
    )
    doc.add_paragraph(
        'Federal Transit Administration. (2024). National Transit Database: '
        '2023 Annual Report. U.S. Department of Transportation.'
    )
    doc.add_paragraph(
        'Nakamura, K., & Hayashi, Y. (2022). Comparative analysis of bus '
        'rapid transit systems in developing megacities. Journal of Transport '
        'Geography, 101, 103-118.'
    )
    doc.add_paragraph(
        'Thompson, L. S. (2023). The economics of transit fleet replacement: '
        'Lifecycle cost models for electric bus adoption. Public Works '
        'Management & Policy, 28(3), 312-330.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
