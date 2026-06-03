"""
Initial Setup: Research paper with single-spaced paragraphs
Task ID: wrpara_013
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_013'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- Title ---
    title = doc.add_heading('', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('The Impact of Urban Green Spaces on Mental Health Outcomes: A Longitudinal Study')
    run.font.size = Pt(16)
    run.bold = True

    # --- Abstract Heading ---
    doc.add_heading('Abstract', level=1)

    # --- Abstract Paragraph (single-spaced) ---
    abstract_text = (
        'Urban green spaces have been increasingly recognized as critical determinants of '
        'public mental health in densely populated metropolitan areas. This longitudinal study '
        'examined the association between proximity to urban parks, community gardens, and '
        'tree-lined corridors and self-reported mental health outcomes among 2,847 adult '
        'residents across five major European cities over a 36-month period from January 2022 '
        'to December 2024. Using validated instruments including the General Health '
        'Questionnaire (GHQ-12) and the Warwick-Edinburgh Mental Wellbeing Scale (WEMWBS), '
        'we found statistically significant improvements in anxiety reduction (p < 0.001), '
        'stress resilience (p = 0.003), and overall life satisfaction (p < 0.01) among '
        'participants residing within 500 meters of a qualifying green space. These findings '
        'suggest that urban planning policies prioritizing accessible green infrastructure '
        'may yield measurable public health benefits, particularly in neighborhoods with '
        'limited recreational amenities.'
    )
    abstract_para = doc.add_paragraph(abstract_text)
    abstract_para.paragraph_format.space_after = Pt(12)

    # --- Introduction Heading ---
    doc.add_heading('Introduction', level=1)

    # --- Body Paragraph 1 ---
    body1_text = (
        'The relationship between natural environments and human psychological wellbeing '
        'has been a subject of academic inquiry since the seminal work of Kaplan and Kaplan '
        'in 1989, which proposed the Attention Restoration Theory as a framework for '
        'understanding how exposure to nature facilitates cognitive recovery from mental '
        'fatigue. Subsequent decades of research have expanded this theoretical foundation '
        'to encompass broader dimensions of mental health, including emotional regulation, '
        'social cohesion, and resilience to chronic psychological stressors prevalent in '
        'urban settings.'
    )
    doc.add_paragraph(body1_text)

    # --- Body Paragraph 2 ---
    body2_text = (
        'Despite a growing body of evidence supporting the mental health benefits of green '
        'space access, significant methodological gaps persist in the existing literature. '
        'Cross-sectional designs dominate the field, limiting causal inference. Additionally, '
        'many studies rely on crude proximity measures that fail to account for the quality, '
        'biodiversity, and maintenance status of green spaces, which recent meta-analyses '
        'by Richardson et al. (2023) have identified as potentially critical moderating variables '
        'affecting health outcomes.'
    )
    doc.add_paragraph(body2_text)

    # --- Body Paragraph 3 ---
    body3_text = (
        'The present study addresses these limitations through a prospective cohort design '
        'that tracked participants across multiple assessment points over three years. We '
        'employed Geographic Information System mapping to characterize green space '
        'availability within 250-meter, 500-meter, and 1-kilometer radii of each '
        "participant's residence, incorporating granular measures of vegetation density, "
        'canopy coverage, and recreational facility presence. Our sampling strategy '
        'intentionally overrepresented lower-income neighborhoods to examine equity '
        'dimensions of green space access and associated mental health disparities.'
    )
    doc.add_paragraph(body3_text)

    # --- Body Paragraph 4 ---
    body4_text = (
        'Our primary hypotheses were threefold: first, that greater proximity to urban green '
        'spaces would predict improved mental health outcomes after controlling for '
        'socioeconomic confounders; second, that this association would strengthen over time '
        'as cumulative exposure increased; and third, that the quality and biodiversity of '
        'accessible green spaces would moderate the strength of the proximity-wellbeing '
        'relationship. These hypotheses were tested using multilevel growth modeling with '
        'random intercepts and slopes, accounting for the nested structure of repeated '
        'measurements within individuals and individuals within neighborhoods.'
    )
    doc.add_paragraph(body4_text)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
