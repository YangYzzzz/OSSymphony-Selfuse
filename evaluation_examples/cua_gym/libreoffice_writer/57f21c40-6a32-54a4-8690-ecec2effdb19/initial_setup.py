"""
Initial Setup: Create a Writer document with 12 bibliography citations throughout,
but no bibliography section at the end.
Task ID: writer_bs_037
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_037'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading('The Impact of Artificial Intelligence on Modern Healthcare Systems', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Author line
    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('Dr. Elena Vasquez, Department of Health Informatics')
    run.font.size = Pt(11)
    run.italic = True

    # Abstract
    doc.add_heading('Abstract', level=1)
    abstract = doc.add_paragraph(
        'The integration of artificial intelligence (AI) into healthcare has accelerated '
        'dramatically over the past decade. This paper examines the multifaceted impact of '
        'AI technologies on diagnostic accuracy, patient outcomes, and operational efficiency '
        'across various healthcare settings. Drawing on recent studies spanning radiology, '
        'pathology, and primary care, we present a comprehensive analysis of both the '
        'opportunities and challenges associated with AI-driven healthcare transformation.'
    )
    abstract.paragraph_format.space_after = Pt(12)

    # Section 1 - Introduction (cites: Chen 2023, Patel 2022, Rodriguez 2024)
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'The rapid advancement of machine learning algorithms has fundamentally altered '
        'the landscape of modern medicine. As Chen (2023) demonstrated in a landmark study '
        'of over 50,000 patient records, AI-assisted diagnostic tools can reduce misdiagnosis '
        'rates by up to 37% in emergency department settings. This finding builds on earlier '
        'work by Patel and Kumar (2022), who first identified the potential for deep learning '
        'models to outperform human radiologists in detecting early-stage pulmonary nodules.'
    )
    doc.add_paragraph(
        'The scope of AI applications in healthcare extends well beyond diagnostic imaging. '
        'Rodriguez et al. (2024) recently published a comprehensive taxonomy of AI use cases '
        'spanning clinical decision support, drug discovery, administrative workflow optimization, '
        'and personalized treatment planning. Their framework provides a useful lens through '
        'which to evaluate the current state of the field.'
    )

    # Section 2 - Diagnostic Applications (cites: Nakamura 2023, Williams 2021, Okonkwo 2023)
    doc.add_heading('2. Diagnostic Applications', level=1)
    doc.add_paragraph(
        'In diagnostic radiology, convolutional neural networks (CNNs) have demonstrated '
        'remarkable proficiency. Nakamura and Tanaka (2023) reported that their ensemble '
        'model achieved 94.2% sensitivity and 97.1% specificity in mammographic screening, '
        'surpassing the average performance of board-certified radiologists. These results '
        'were corroborated by Williams (2021), whose meta-analysis of 23 studies confirmed '
        'a consistent advantage for AI-assisted readings in breast cancer detection.'
    )
    doc.add_paragraph(
        'Beyond imaging, AI has shown promise in pathology and genomics. Okonkwo and '
        'Adeyemi (2023) developed a transformer-based model capable of predicting tumor '
        'microsatellite instability status from standard hematoxylin and eosin-stained slides '
        'with an AUC of 0.91. This approach could significantly reduce the need for costly '
        'molecular testing in colorectal cancer staging.'
    )

    # Section 3 - Patient Outcomes (cites: Johansson 2022, Kim 2024, Al-Rashidi 2023)
    doc.add_heading('3. Impact on Patient Outcomes', level=1)
    doc.add_paragraph(
        'The translation of AI capabilities into measurable patient benefits has been a '
        'central concern. Johansson et al. (2022) conducted a randomized controlled trial '
        'across 14 Swedish hospitals, finding that AI-augmented triage reduced average '
        'emergency department wait times by 22 minutes and decreased 30-day readmission '
        'rates by 8.4 percentage points. Kim and Park (2024) extended this work to the '
        'intensive care setting, demonstrating that predictive analytics for sepsis onset '
        'improved survival rates by 12% when integrated into nursing workflows.'
    )
    doc.add_paragraph(
        'However, not all outcomes have been uniformly positive. Al-Rashidi (2023) '
        'cautioned that algorithmic bias in training datasets can lead to disparate '
        'performance across demographic groups. In a study of dermatological AI tools, '
        'diagnostic accuracy for skin lesions was 18% lower for patients with darker '
        'skin tones, highlighting the urgent need for diverse and representative training data.'
    )

    # Section 4 - Operational Efficiency (cites: Fernandez 2022, Gupta 2023)
    doc.add_heading('4. Operational Efficiency', level=1)
    doc.add_paragraph(
        'Administrative applications of AI have yielded substantial cost savings. '
        'Fernandez and Morales (2022) reported that natural language processing (NLP) tools '
        'for clinical documentation reduced physician charting time by an average of 45 minutes '
        'per shift, translating to an estimated annual savings of $1.2 million per 500-bed '
        'hospital. Gupta et al. (2023) found similar efficiencies in claims processing, where '
        'AI-driven prior authorization systems decreased approval turnaround from 72 hours to '
        'under 4 hours while maintaining accuracy rates above 98%.'
    )

    # Section 5 - Conclusion (cites: Liu 2024)
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        'The evidence presented in this review suggests that AI has the potential to '
        'transform healthcare delivery across multiple dimensions. However, as Liu (2024) '
        'emphasized in a recent editorial, the successful integration of AI into clinical '
        'practice requires careful attention to ethical considerations, regulatory frameworks, '
        'and the preservation of the physician-patient relationship. Future research should '
        'prioritize longitudinal studies that track both the clinical and socioeconomic '
        'impacts of AI adoption in diverse healthcare systems.'
    )

    # NOTE: No bibliography section - that is the task the agent must complete

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
