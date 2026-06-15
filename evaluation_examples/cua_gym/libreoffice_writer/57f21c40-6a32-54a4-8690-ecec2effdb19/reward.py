"""
Reward Script: Create bibliography section with two columns, 'References' heading, and hanging indent
Task ID: writer_bs_037
Domain: libreoffice_writer
Scoring:
  Component 1: 'References' heading in Heading 1 style (0.20)
  Component 2: Bibliography section in two-column layout (0.20)
  Component 3: Entries have hanging indent of 1.27cm (0.25)
  Component 4: At least 12 bibliography entries present (0.20)
  Component 5: Entries alphabetically ordered by author (0.15)
"""

import os
from docx import Document
from docx.shared import Emu
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_037'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        import time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the 'References' heading paragraph index
    refs_idx = None
    refs_style = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() == 'references':
            refs_idx = i
            refs_style = p.style.name
            break

    # Component 1: 'References' heading in Heading 1 style (0.20 points)
    try:
        if refs_idx is not None and refs_style == 'Heading 1':
            print(f"PASS: Component 1 — 'References' heading found at para {refs_idx} with style 'Heading 1' (0.20 pts)")
            total_score += 0.20
        elif refs_idx is not None:
            print(f"FAIL: Component 1 — 'References' heading found but style is '{refs_style}', expected 'Heading 1'")
        else:
            print("FAIL: Component 1 — No 'References' heading found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # If no References heading found, remaining checks are moot
    if refs_idx is None:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Component 2: Bibliography section uses two-column layout (0.20 points)
    # The golden document has a new continuous section break before the bibliography
    # with w:cols w:num="2"
    try:
        two_col_count = sum(
            1 for sec in doc.sections
            if sec._sectPr.find(qn('w:cols')) is not None
            and sec._sectPr.find(qn('w:cols')).get(qn('w:num')) == '2'
        )

        if two_col_count > 0:
            print(f"PASS: Component 2 — Two-column section found (0.20 pts)")
            total_score += 0.20
        else:
            print("FAIL: Component 2 — No two-column section found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Collect bibliography entries (non-empty paragraphs after the References heading)
    bib_entries = []
    for p in doc.paragraphs[refs_idx + 1:]:
        if p.text.strip():
            bib_entries.append(p)

    # Component 3: Entries have hanging indent (left_indent=1.27cm, first_line=-1.27cm) (0.25 points)
    try:
        if len(bib_entries) == 0:
            print("FAIL: Component 3 — No bibliography entries found to check indentation")
        else:
            # 1.27cm = 457200 EMU (1 cm = 360000 EMU, so 1.27 * 360000 = 457200)
            target_indent_emu = 457200
            tolerance_emu = 20000  # small tolerance for rounding

            correct_indent_count = 0
            for entry in bib_entries:
                pf = entry.paragraph_format
                li = pf.left_indent
                fli = pf.first_line_indent

                if li is not None and fli is not None:
                    li_val = int(li)
                    fli_val = int(fli)
                    # Hanging indent: left_indent ~ 1.27cm, first_line_indent ~ -1.27cm
                    if (abs(li_val - target_indent_emu) <= tolerance_emu and
                            abs(fli_val - (-target_indent_emu)) <= tolerance_emu):
                        correct_indent_count += 1

            ratio = correct_indent_count / len(bib_entries)
            if ratio >= 0.9:
                print(f"PASS: Component 3 — {correct_indent_count}/{len(bib_entries)} entries have correct hanging indent (0.25 pts)")
                total_score += 0.25
            elif ratio >= 0.5:
                partial = round(0.25 * ratio, 2)
                print(f"PARTIAL: Component 3 — {correct_indent_count}/{len(bib_entries)} entries have correct hanging indent ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — Only {correct_indent_count}/{len(bib_entries)} entries have correct hanging indent")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: At least 12 bibliography entries present (0.20 points)
    try:
        num_entries = len(bib_entries)
        if num_entries >= 12:
            print(f"PASS: Component 4 — {num_entries} bibliography entries found (>= 12) (0.20 pts)")
            total_score += 0.20
        elif num_entries >= 6:
            partial = round(0.20 * (num_entries / 12), 2)
            print(f"PARTIAL: Component 4 — {num_entries} entries found (expected >= 12) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {num_entries} bibliography entries found (expected >= 12)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Entries alphabetically ordered by author last name (0.15 points)
    try:
        if len(bib_entries) < 2:
            print("FAIL: Component 5 — Not enough entries to check alphabetical order")
        else:
            # Extract first author's last name from each entry
            authors = []
            for entry in bib_entries:
                text = entry.text.strip()
                # Typical format: "LastName, F. (YYYY). Title..."
                # Extract the part before the first '('
                author_part = text.split('(')[0].strip() if '(' in text else text
                # Get first author's last name (before first comma)
                last_name = author_part.split(',')[0].strip().lower()
                authors.append(last_name)

            is_sorted = all(authors[i] <= authors[i + 1] for i in range(len(authors) - 1))
            if is_sorted:
                print(f"PASS: Component 5 — Entries are alphabetically ordered by author ({len(authors)} entries) (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 5 — Entries are NOT alphabetically ordered: {authors}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
