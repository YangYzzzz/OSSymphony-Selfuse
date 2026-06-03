"""
Reward Script: Information Security Policy Document
Task ID: writer_wf_038
Domain: libreoffice_writer
Scoring:
  Component 1: Title present (0.10)
  Component 2: Version control table (0.15)
  Component 3: Six Heading 1 sections (0.25)
  Component 4: Five numbered sub-policies (0.15)
  Component 5: Roles table (0.15)
  Component 6: Confidentiality footer (0.20)
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_038'


def persist_app_state(domain: str):
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect paragraph info
    paragraphs = doc.paragraphs
    tables = doc.tables

    # Component 1: Title "Information Security Policy" (0.10 points)
    try:
        has_title = False
        for p in paragraphs:
            if p.style and p.style.name in ('Title', 'Heading 0'):
                if 'information security policy' in p.text.lower():
                    has_title = True
                    break
        if has_title:
            print(f"PASS: Component 1 — Title 'Information Security Policy' found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — Title 'Information Security Policy' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Version control table with correct structure (0.15 points)
    # Must have a table with 4 columns (Version, Date, Author, Changes) and at least 4 rows (header + 3 versions)
    try:
        version_table_found = False
        for table in tables:
            if len(table.columns) >= 4 and len(table.rows) >= 4:
                header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
                if 'version' in header_cells and 'date' in header_cells:
                    # Check that at least 3 data rows exist (non-empty first cell)
                    data_rows = 0
                    for ri in range(1, len(table.rows)):
                        first_cell = table.rows[ri].cells[0].text.strip()
                        if first_cell:
                            data_rows += 1
                    if data_rows >= 3:
                        version_table_found = True
                        break
        if version_table_found:
            print(f"PASS: Component 2 — Version control table found with >=3 data rows (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 — Version control table not found or insufficient rows")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Six required Heading 1 sections (0.25 points)
    # Required: Purpose, Scope, Policy Statements, Roles and Responsibilities, Compliance, Review and Updates
    try:
        required_sections = [
            'purpose',
            'scope',
            'policy statements',
            'roles and responsibilities',
            'compliance',
            'review and updates',
        ]
        heading1_texts = []
        for p in paragraphs:
            if p.style and p.style.name == 'Heading 1':
                heading1_texts.append(p.text.strip().lower())

        found_sections = 0
        for req in required_sections:
            if any(req in h for h in heading1_texts):
                found_sections += 1

        if found_sections == 6:
            print(f"PASS: Component 3 — All 6 Heading 1 sections found (0.25 pts)")
            total_score += 0.25
        elif found_sections >= 4:
            partial = round(0.25 * (found_sections / 6), 2)
            print(f"PARTIAL: Component 3 — {found_sections}/6 sections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {found_sections}/6 required sections found. Found headings: {heading1_texts}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Five numbered sub-policies under Policy Statements (0.15 points)
    # Looking for paragraphs starting with "1.", "2.", "3.", "4.", "5." between Policy Statements and the next heading
    try:
        # Find the index range for Policy Statements section
        ps_start = None
        ps_end = None
        for i, p in enumerate(paragraphs):
            if p.style and p.style.name == 'Heading 1' and 'policy statements' in p.text.lower():
                ps_start = i
            elif ps_start is not None and p.style and p.style.name == 'Heading 1':
                ps_end = i
                break
        if ps_end is None:
            ps_end = len(paragraphs)

        numbered_policies = 0
        if ps_start is not None:
            for i in range(ps_start + 1, ps_end):
                text = paragraphs[i].text.strip()
                # Check if starts with a number followed by dot/period
                if re.match(r'^\d+[\.\)]\s', text):
                    numbered_policies += 1

        if numbered_policies >= 5:
            print(f"PASS: Component 4 — {numbered_policies} numbered sub-policies found (0.15 pts)")
            total_score += 0.15
        elif numbered_policies >= 3:
            partial = round(0.15 * (numbered_policies / 5), 2)
            print(f"PARTIAL: Component 4 — {numbered_policies}/5 sub-policies ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {numbered_policies} numbered sub-policies found (need >= 5)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Roles and Responsibilities table (0.15 points)
    # Must have a table with 2 columns (Role, Responsibility) and at least 5 rows (header + 4 roles)
    try:
        roles_table_found = False
        for table in tables:
            if len(table.columns) >= 2:
                header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
                if 'role' in header_cells and 'responsibility' in header_cells:
                    # Check for at least 4 data rows
                    data_rows = 0
                    for ri in range(1, len(table.rows)):
                        first_cell = table.rows[ri].cells[0].text.strip()
                        if first_cell:
                            data_rows += 1
                    if data_rows >= 4:
                        roles_table_found = True
                        break
        if roles_table_found:
            print(f"PASS: Component 5 — Roles table found with >=4 data rows (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — Roles and Responsibilities table not found or insufficient rows")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Confidentiality notice in footer (0.20 points)
    try:
        footer_found = False
        for section in doc.sections:
            footer = section.footer
            if footer and footer.paragraphs:
                footer_text = ' '.join(p.text for p in footer.paragraphs).strip()
                if 'confidential' in footer_text.lower() and 'internal use only' in footer_text.lower():
                    footer_found = True
                    break
        if footer_found:
            print(f"PASS: Component 6 — Confidentiality footer found (0.20 pts)")
            total_score += 0.20
        else:
            # Check all sections
            all_footer_text = []
            for section in doc.sections:
                if section.footer and section.footer.paragraphs:
                    all_footer_text.append(' '.join(p.text for p in section.footer.paragraphs))
            print(f"FAIL: Component 6 — Footer text: {all_footer_text}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist app state before verification
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
