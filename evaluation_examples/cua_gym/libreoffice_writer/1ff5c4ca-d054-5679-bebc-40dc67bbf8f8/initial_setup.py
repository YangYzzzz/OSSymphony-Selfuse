"""
Initial Setup: Create a Writer document with 20 bibliography entries, no special indentation.
Task ID: writer_fs_046
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_046'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("References", level=1)

    # 20 realistic bibliography entries (APA-style)
    references = [
        "Anderson, R. C., & Pearson, P. D. (2024). A schema-theoretic view of basic processes in reading comprehension. In P. D. Pearson (Ed.), Handbook of reading research (pp. 255-291). Longman.",
        "Baker, L., & Brown, A. L. (2023). Metacognitive skills and reading. In P. D. Pearson (Ed.), Handbook of reading research (pp. 353-394). Longman.",
        "Chen, W., & Liu, Y. (2025). Machine learning approaches to automated essay scoring: A comprehensive review. Journal of Educational Technology, 48(3), 112-134.",
        "Davidson, K. M., & Thompson, R. J. (2022). The role of emotional regulation in academic performance among college students. Educational Psychology Review, 34(2), 567-589.",
        "Ellis, R., & Shintani, N. (2023). Exploring language pedagogy through second language acquisition research. Routledge.",
        "Fernandez, M. A., Garcia, J. L., & Rodriguez, P. (2024). Digital literacy in higher education: A systematic literature review. Computers & Education, 189, 104582.",
        "Gonzalez, T., & Martinez, S. (2023). Impact of remote learning on student engagement during the post-pandemic era. International Journal of Educational Research, 117, 102-118.",
        "Harrison, C., & Killion, J. (2022). Ten roles for teacher leaders: Professional development strategies for school improvement. Educational Leadership, 65(1), 74-77.",
        "Ibrahim, N., & Shahrill, M. (2025). A systematic review on the effectiveness of flipped classroom approach in mathematics education. Mathematics Education Research Journal, 37(1), 45-67.",
        "Jensen, L., & Konradsen, F. (2024). A review of the use of virtual reality head-mounted displays in education and training. Education and Information Technologies, 23(4), 1515-1529.",
        "Kim, S., & Park, H. (2023). The influence of collaborative learning on critical thinking skills in online environments. The Internet and Higher Education, 56, 100-115.",
        "Lambert, W. E., & Tucker, G. R. (2022). Bilingual education of children: The St. Lambert experiment. Newbury House.",
        "Mitchell, R., Myles, F., & Marsden, E. (2024). Second language learning theories (4th ed.). Routledge.",
        "Nguyen, T. H., & Williams, A. (2025). Artificial intelligence in personalized learning: Current trends and future directions. British Journal of Educational Technology, 56(2), 234-251.",
        "O'Brien, D. G., & Stewart, R. A. (2023). Resistance to content area reading instruction: Dimensions and solutions. In E. K. Dishner (Ed.), Reading in the content areas (3rd ed., pp. 30-40). Kendall/Hunt.",
        "Patel, R., & Singh, V. (2024). Gamification in education: A systematic mapping study. Educational Technology & Society, 27(3), 45-63.",
        "Quinn, J., & Spencer, M. (2022). Exploring the relationship between socioeconomic status and academic achievement: A meta-analysis. Review of Educational Research, 92(4), 612-645.",
        "Ramirez, G., Chang, H., & Maloney, E. A. (2023). Math anxiety: Past research, promising interventions, and a new interpretation framework. Educational Psychologist, 51(3-4), 404-425.",
        "Stevens, L. P., & Bean, T. W. (2024). Critical literacy: Context, research, and practice in the K-12 classroom (2nd ed.). SAGE Publications.",
        "Torres, A., Dominguez, C., & Valdes, R. (2025). Blockchain technology for academic credential verification: Opportunities and challenges. Journal of Higher Education Policy and Management, 47(1), 89-104.",
    ]

    for ref in references:
        para = doc.add_paragraph(ref)
        # Ensure no special indentation
        para.paragraph_format.left_indent = None
        para.paragraph_format.first_line_indent = None
        # Set a readable font
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
