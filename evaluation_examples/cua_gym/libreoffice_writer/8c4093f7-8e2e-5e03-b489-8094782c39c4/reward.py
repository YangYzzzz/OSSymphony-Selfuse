"""
Reward Script: Operations Manual in LibreOffice Writer
Task ID: writer_wf_028
Domain: libreoffice_writer
Scoring:
  Component 1: Title page with correct title (0.15)
  Component 2: 5 chapter headings (Heading 1) with correct names (0.25)
  Component 3: 10 subsections (Heading 2) (0.20)
  Component 4: Subsection body content (0.15)
  Component 5: Page breaks before chapters (0.15)
  Component 6: Page numbering in footer (0.10)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_028'

EXPECTED_CHAPTERS = [
    'Receiving Procedures',
    'Inventory Management',
    'Order Fulfillment',
    'Shipping Protocols',
    'Safety Guidelines',
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have content (blank doc = initial state)
    if len(doc.paragraphs) < 5:
        print("FAIL: Document has fewer than 5 paragraphs — appears blank or nearly empty")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraphs with their styles
    all_paras = [(p.style.name if p.style else 'None', p.text.strip()) for p in doc.paragraphs]

    # Component 1: Title page with correct title text (0.15 points)
    try:
        title_found = False
        for style, text in all_paras:
            if 'operations manual' in text.lower() and 'distribution center' in text.lower():
                title_found = True
                break
        if title_found:
            print(f"PASS: Component 1 — Title 'Operations Manual - Distribution Center' found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Title with 'Operations Manual' and 'Distribution Center' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 5 chapter headings (Heading 1) with correct names (0.25 points)
    # Exclude TOC heading — only count the 5 actual chapter headings
    try:
        h1_paras = [text for style, text in all_paras if style == 'Heading 1' and text]
        # Match expected chapters (case-insensitive, partial match)
        matched_chapters = 0
        for expected in EXPECTED_CHAPTERS:
            for h1 in h1_paras:
                if expected.lower() in h1.lower():
                    matched_chapters += 1
                    break
        if matched_chapters == 5:
            print(f"PASS: Component 2 — All 5 chapter headings found as Heading 1 (0.25 pts)")
            total_score += 0.25
        elif matched_chapters >= 3:
            partial = round(0.25 * (matched_chapters / 5), 2)
            print(f"PARTIAL: Component 2 — {matched_chapters}/5 chapter headings found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {matched_chapters}/5 chapter headings found. H1s: {h1_paras}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: 10 subsections with Heading 2 style (0.20 points)
    try:
        h2_paras = [text for style, text in all_paras if style == 'Heading 2' and text]
        h2_count = len(h2_paras)
        if h2_count >= 10:
            print(f"PASS: Component 3 — {h2_count} Heading 2 subsections found (>=10) (0.20 pts)")
            total_score += 0.20
        elif h2_count >= 5:
            partial = round(0.20 * (h2_count / 10), 2)
            print(f"PARTIAL: Component 3 — {h2_count}/10 subsections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {h2_count} Heading 2 subsections found, expected 10")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Subsection body content — each Heading 2 should be followed by at least
    # one Normal paragraph with substantial text (2+ sentences ~ 40+ chars) (0.15 points)
    try:
        body_count = 0
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name == 'Heading 2':
                # Look at the next paragraph(s) for body text
                for j in range(i + 1, min(i + 4, len(doc.paragraphs))):
                    next_p = doc.paragraphs[j]
                    next_style = next_p.style.name if next_p.style else ''
                    if next_style.startswith('Heading'):
                        break
                    if next_p.text.strip() and len(next_p.text.strip()) >= 40:
                        body_count += 1
                        break
        if body_count >= 10:
            print(f"PASS: Component 4 — {body_count}/10 subsections have body content (0.15 pts)")
            total_score += 0.15
        elif body_count >= 5:
            partial = round(0.15 * (body_count / 10), 2)
            print(f"PARTIAL: Component 4 — {body_count}/10 subsections have body content ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {body_count}/10 subsections have body content")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Page breaks before chapters (0.15 points)
    # Each of the 5 chapter Heading 1s should be preceded by a page break
    # (either inline in a prior paragraph or page_break_before on the heading)
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        chapter_breaks = 0

        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name == 'Heading 1':
                text = para.text.strip().lower()
                # Only check the 5 actual chapters, not TOC
                is_chapter = any(ch.lower() in text for ch in EXPECTED_CHAPTERS)
                if not is_chapter:
                    continue

                has_break = False
                # Check page_break_before on the heading itself
                pPr = para._element.find('.//w:pPr', ns)
                if pPr is not None:
                    pgBB = pPr.find('w:pageBreakBefore', ns)
                    if pgBB is not None:
                        has_break = True

                # Check inline page break in the preceding paragraph(s)
                if not has_break and i > 0:
                    for k in range(max(0, i - 2), i):
                        prev_para = doc.paragraphs[k]
                        for run in prev_para.runs:
                            for br in run.element.findall('.//w:br', ns):
                                btype = br.attrib.get(
                                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', '')
                                if btype == 'page':
                                    has_break = True
                                    break
                            if has_break:
                                break
                        if has_break:
                            break

                if has_break:
                    chapter_breaks += 1

        if chapter_breaks >= 5:
            print(f"PASS: Component 5 — Page breaks found before all 5 chapters (0.15 pts)")
            total_score += 0.15
        elif chapter_breaks >= 3:
            partial = round(0.15 * (chapter_breaks / 5), 2)
            print(f"PARTIAL: Component 5 — Page breaks before {chapter_breaks}/5 chapters ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Page breaks before only {chapter_breaks}/5 chapters")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Page numbering in footer (0.10 points)
    # Footer should contain a PAGE field code
    try:
        has_page_number = False
        for section in doc.sections:
            footer = section.footer
            if footer and footer.paragraphs:
                for fp in footer.paragraphs:
                    xml_str = fp._element.xml
                    if 'PAGE' in xml_str and 'instrText' in xml_str:
                        has_page_number = True
                        break
                    # Also check for fldChar pattern
                    if 'fldChar' in xml_str:
                        has_page_number = True
                        break
            if has_page_number:
                break

        if has_page_number:
            print(f"PASS: Component 6 — Page numbering found in footer (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 — No page numbering field found in any footer")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint: persist app state then verify
def persist_app_state(domain):
    import os
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
