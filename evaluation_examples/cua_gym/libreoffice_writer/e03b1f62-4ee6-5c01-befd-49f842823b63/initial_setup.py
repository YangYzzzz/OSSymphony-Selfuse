"""
Initial Setup: Three-column newsletter with 'Local Events' near bottom of second column.
Task ID: writer_fs_033
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_033'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_paragraph(doc, text, level=1):
    """Add a heading with newsletter styling."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if level == 0:
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.paragraph_format.space_after = Pt(4)
    elif level == 1:
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
    return para


def add_body_text(doc, text):
    """Add body text with newsletter styling."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return para


def set_columns(section, num_cols, spacing_inches=0.3):
    """Set the number of columns for a section using XML manipulation."""
    sectPr = section._sectPr
    cols_elem = sectPr.find(qn('w:cols'))
    if cols_elem is None:
        cols_elem = sectPr.makeelement(qn('w:cols'), {})
        sectPr.append(cols_elem)
    cols_elem.set(qn('w:num'), str(num_cols))
    cols_elem.set(qn('w:space'), str(int(spacing_inches * 1440)))  # inches to twips


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Set 3-column layout
    set_columns(section, 3, spacing_inches=0.3)

    # Newsletter title
    add_heading_paragraph(doc, 'The Riverside Gazette', level=0)

    # Subtitle / dateline
    sub = doc.add_paragraph()
    run = sub.add_run('Volume 12, Issue 4  |  March 2025  |  Serving the Riverside Community Since 2013')
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sub.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.paragraph_format.space_after = Pt(8)

    # Separator line
    sep = doc.add_paragraph()
    run = sep.add_run('_' * 90)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    sep.paragraph_format.space_after = Pt(8)

    # --- Article 1: New Park Development ---
    add_heading_paragraph(doc, 'New Park Development Approved')
    add_body_text(doc,
        'The Riverside City Council voted unanimously on Tuesday evening to approve '
        'the construction of a new 15-acre community park on the former industrial '
        'site along Elm Street. The project, estimated at $4.2 million, will include '
        'walking trails, a playground, and a community garden.')
    add_body_text(doc,
        'Mayor Patricia Hernandez praised the decision, calling it "a transformative '
        'investment in our community\'s future." Construction is expected to begin in '
        'June 2025 and be completed by late 2026.')
    add_body_text(doc,
        'Residents in the nearby Oak Hill neighborhood have expressed overwhelming '
        'support for the project, with over 800 signatures collected in a petition '
        'presented at the council meeting. The park will also feature a small amphitheater '
        'for outdoor performances and community gatherings.')

    # --- Article 2: School Board Update ---
    add_heading_paragraph(doc, 'School Board Announces Budget Increase')
    add_body_text(doc,
        'The Riverside School Board has approved a 7.3% increase in the annual '
        'education budget, bringing total funding to $28.6 million for the 2025-2026 '
        'academic year. Superintendent Dr. James Watkins outlined key priorities '
        'including technology upgrades, teacher retention bonuses, and expanded '
        'after-school programs.')
    add_body_text(doc,
        'The budget increase will fund 12 new teaching positions across three elementary '
        'schools and the addition of a STEM lab at Riverside Middle School. Board member '
        'Angela Torres noted that reading scores improved 15% this year thanks to the '
        'literacy initiative launched in September 2024.')
    add_body_text(doc,
        'Parent organizations have welcomed the announcement, though some have called '
        'for greater transparency in how funds are allocated across the district. A public '
        'forum is scheduled for April 10th at the community center to address questions.')

    # --- Article 3: Farmers Market Season ---
    add_heading_paragraph(doc, 'Farmers Market Opens for Spring Season')
    add_body_text(doc,
        'The beloved Riverside Farmers Market returns this Saturday at Centennial Plaza, '
        'marking the beginning of its 9th season. Over 40 local vendors will offer fresh '
        'produce, artisan breads, handmade cheeses, and organic honey.')
    add_body_text(doc,
        'Market organizer Sarah Chen expects record attendance this year. "We have '
        'eight new vendors joining us, including a family-run flower farm from Willow '
        'Creek and a craft brewery offering tastings," she said. The market runs every '
        'Saturday from 8 AM to 1 PM through October.')

    # --- Article 4: Local Events (this should land near bottom of column 2) ---
    add_heading_paragraph(doc, 'Local Events')
    add_body_text(doc,
        'April 5 - Spring Art Walk: Downtown galleries open their doors for the annual '
        'art walk from 5 PM to 9 PM. Over 30 artists will showcase paintings, sculptures, '
        'and photography. Refreshments provided by local restaurants.')
    add_body_text(doc,
        'April 12 - Community Clean-Up Day: Volunteers needed for the annual neighborhood '
        'beautification project. Meet at City Hall at 9 AM. Supplies and lunch provided. '
        'Last year, over 200 volunteers collected 3 tons of litter and planted 150 trees.')
    add_body_text(doc,
        'April 19 - Riverside 5K Fun Run: Registration is open for the charity run '
        'benefiting the Riverside Children\'s Hospital. Entry fee is $25 for adults and '
        '$10 for children under 12. The route follows the scenic river trail.')
    add_body_text(doc,
        'April 26 - Book Fair at the Library: The Riverside Public Library hosts its '
        'semi-annual book fair with thousands of donated books at bargain prices. All '
        'proceeds support library programs and summer reading initiatives.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
