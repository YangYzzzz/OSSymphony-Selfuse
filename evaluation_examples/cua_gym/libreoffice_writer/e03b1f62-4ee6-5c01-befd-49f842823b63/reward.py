"""
Reward Script: Insert column break before 'Local Events' heading
Task ID: writer_fs_033
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Column break exists in the 'Local Events' paragraph
  Component 2 (0.3): Column break is positioned before the 'Local Events' text
  Component 3 (0.3): Document text integrity - 'Local Events' and content preserved
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_033'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def find_local_events_para(doc):
    """Find the paragraph containing 'Local Events' heading text."""
    for i, para in enumerate(doc.paragraphs):
        if 'Local Events' in para.text:
            return i, para
    return None, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the 'Local Events' paragraph
    para_idx, le_para = find_local_events_para(doc)
    if le_para is None:
        print("CRITICAL: 'Local Events' paragraph not found in document")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Found 'Local Events' at paragraph index {para_idx}")

    # Track whether column break was found (used as gate for later components)
    column_break_found = False

    # Component 1: Column break exists in the 'Local Events' paragraph (0.4 points)
    try:
        for run in le_para.runs:
            for br in run.element.findall(f'.//{{{W_NS}}}br'):
                br_type = br.get(f'{{{W_NS}}}type', '')
                if br_type == 'column':
                    column_break_found = True
                    break
            if column_break_found:
                break

        # Also check via raw XML children of the paragraph element for breaks
        # in runs that python-docx may not expose via .runs
        if not column_break_found:
            for r_elem in le_para._element.findall(f'{{{W_NS}}}r'):
                for br in r_elem.findall(f'{{{W_NS}}}br'):
                    br_type = br.get(f'{{{W_NS}}}type', '')
                    if br_type == 'column':
                        column_break_found = True
                        break
                if column_break_found:
                    break

        if column_break_found:
            print(f"PASS: Component 1 - Column break found in 'Local Events' paragraph (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 - No column break in 'Local Events' paragraph")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Column break is positioned BEFORE the 'Local Events' text (0.3 points)
    # The break run should come before the text run in the paragraph XML
    try:
        break_position = None
        text_position = None
        for idx, r_elem in enumerate(le_para._element.findall(f'{{{W_NS}}}r')):
            # Check if this run has a column break
            for br in r_elem.findall(f'{{{W_NS}}}br'):
                br_type = br.get(f'{{{W_NS}}}type', '')
                if br_type == 'column' and break_position is None:
                    break_position = idx
            # Check if this run has the 'Local Events' text
            for t_elem in r_elem.findall(f'{{{W_NS}}}t'):
                if t_elem.text and 'Local Events' in t_elem.text and text_position is None:
                    text_position = idx

        if break_position is not None and text_position is not None:
            if break_position < text_position:
                print(f"PASS: Component 2 - Column break (run {break_position}) precedes text (run {text_position}) (0.3 pts)")
                total_score += 0.3
            elif break_position == text_position:
                # Break and text in same run - check element order within run
                r_elem = le_para._element.findall(f'{{{W_NS}}}r')[break_position]
                children = list(r_elem)
                br_idx = None
                t_idx = None
                for ci, child in enumerate(children):
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag == 'br' and br_idx is None:
                        if child.get(f'{{{W_NS}}}type', '') == 'column':
                            br_idx = ci
                    if tag == 't' and t_idx is None:
                        if child.text and 'Local Events' in child.text:
                            t_idx = ci
                if br_idx is not None and t_idx is not None and br_idx < t_idx:
                    print(f"PASS: Component 2 - Column break precedes text within same run (0.3 pts)")
                    total_score += 0.3
                else:
                    print(f"FAIL: Component 2 - Column break does not precede 'Local Events' text in run order")
            else:
                print(f"FAIL: Component 2 - Column break (run {break_position}) comes AFTER text (run {text_position})")
        elif break_position is None:
            print(f"FAIL: Component 2 - No column break found to check position")
        else:
            print(f"FAIL: Component 2 - 'Local Events' text run not found for position check")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Column break present AND document text integrity preserved (0.3 points)
    # This is a compound check: column break must exist (task change) AND content must be intact.
    # Gated on column_break_found to ensure initial_env scores 0.0.
    try:
        if not column_break_found:
            print(f"FAIL: Component 3 - Column break not present, skipping content integrity check")
        else:
            all_text = '\n'.join(p.text for p in doc.paragraphs)
            expected_phrases = [
                'Local Events',
                'Spring Art Walk',
                'Community Clean-Up Day',
                'Riverside 5K Fun Run',
                'Book Fair at the Library',
            ]
            found_count = sum(1 for phrase in expected_phrases if phrase in all_text)

            if found_count == len(expected_phrases):
                print(f"PASS: Component 3 - Column break present AND all {len(expected_phrases)} content phrases preserved (0.3 pts)")
                total_score += 0.3
            else:
                partial = round(0.3 * found_count / len(expected_phrases), 2)
                print(f"FAIL: Component 3 - Only {found_count}/{len(expected_phrases)} phrases found (partial: {partial} pts)")
                total_score += partial
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
