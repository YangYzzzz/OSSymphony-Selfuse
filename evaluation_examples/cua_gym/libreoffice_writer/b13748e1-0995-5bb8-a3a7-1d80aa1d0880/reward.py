"""
Reward Script: Insert horizontal line after document header section
Task ID: writer_tech_008
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): A paragraph border (horizontal line) exists in the document
  Component 2 (0.3): The border paragraph is positioned between header block and body content
  Component 3 (0.3): The border has valid properties (single line style, visible)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_008'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def find_border_paragraphs(doc):
    """Find all paragraphs that have any paragraph border (pBdr) element.
    Returns list of (index, paragraph, border_element) tuples.
    """
    results = []
    for i, para in enumerate(doc.paragraphs):
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                # Check that pBdr has at least one actual border child (top, bottom, left, right)
                border_sides = ['w:top', 'w:bottom', 'w:left', 'w:right']
                has_real_border = any(
                    pBdr.find(qn(side)) is not None for side in border_sides
                )
                if has_real_border:
                    results.append((i, para, pBdr))
    return results


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_paras = len(doc.paragraphs)
    print(f"INFO: Document has {num_paras} paragraphs")

    # Identify header block: paragraphs 0-2 should be title, version, date
    # Body starts after header. We expect a separator between them.
    header_texts = []
    for i in range(min(3, num_paras)):
        header_texts.append(doc.paragraphs[i].text.strip())
    print(f"INFO: Header paragraphs: {header_texts}")

    # Find paragraphs with borders
    border_paras = find_border_paragraphs(doc)
    print(f"INFO: Found {len(border_paras)} paragraph(s) with borders")

    # Component 1: A paragraph border (horizontal line) exists in the document (0.4 points)
    # This FAILS on initial (no borders) and PASSES on golden (has border)
    try:
        if len(border_paras) > 0:
            idx, para, pBdr = border_paras[0]
            print(f"PASS: Component 1 - Found paragraph border at index {idx} (0.4 pts)")
            total_score += 0.4
        else:
            print("FAIL: Component 1 - No paragraph with border found in document")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: The border paragraph is positioned between header and body (0.3 points)
    # Header block is paragraphs 0-2 (title, version, date).
    # The border should be at index 3 (right after header, before body content).
    # We check: border paragraph index is 3, and the paragraph after it starts body content
    # (e.g., starts with a section number like "1." or has substantive text).
    # This FAILS on initial (no border paragraph) and PASSES on golden.
    try:
        if len(border_paras) > 0:
            idx, para, pBdr = border_paras[0]
            # The border para should be between header (indices 0-2) and body
            # Accept index 3 (inserted after header) or index 2 (bottom border on last header para)
            if 2 <= idx <= 4:
                # Verify body content follows after the border paragraph
                body_start_idx = idx + 1
                if body_start_idx < num_paras:
                    next_text = doc.paragraphs[body_start_idx].text.strip()
                    # Body content should start with section heading or substantive text
                    if len(next_text) > 0:
                        print(f"PASS: Component 2 - Border at index {idx}, body starts at {body_start_idx} with '{next_text[:50]}' (0.3 pts)")
                        total_score += 0.3
                    else:
                        print(f"FAIL: Component 2 - Border at index {idx} but next paragraph is empty")
                else:
                    print(f"FAIL: Component 2 - Border at index {idx} but no content follows")
            else:
                print(f"FAIL: Component 2 - Border at index {idx}, expected between indices 2-4 (after header)")
        else:
            print("FAIL: Component 2 - No border paragraph found (depends on Component 1)")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: The border has valid properties (0.3 points)
    # Check that the border is a visible line (has val attribute like 'single', 'thick', etc.)
    # and has a reasonable size (sz > 0) and a color.
    # This FAILS on initial (no border) and PASSES on golden.
    try:
        if len(border_paras) > 0:
            idx, para, pBdr = border_paras[0]
            # Check for any border side with valid attributes
            valid_sides = 0
            for side_name in ['w:bottom', 'w:top', 'w:left', 'w:right']:
                side = pBdr.find(qn(side_name))
                if side is not None:
                    val = side.get(qn('w:val'))
                    sz = side.get(qn('w:sz'))
                    color = side.get(qn('w:color'))
                    print(f"INFO: Border side '{side_name}': val={val}, sz={sz}, color={color}")
                    # Valid if val is not 'none' and not 'nil', and sz > 0
                    if val and val not in ('none', 'nil'):
                        if sz is None or int(sz) > 0:
                            valid_sides += 1

            if valid_sides > 0:
                print(f"PASS: Component 3 - Border has valid line properties (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 3 - Border exists but has no valid line style")
        else:
            print("FAIL: Component 3 - No border paragraph found (depends on Component 1)")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
