"""
Initial Setup: Insert a horizontal line after the document header section
Task ID: writer_tech_008
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ===== HEADER BLOCK =====

    # Document Title
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    title_run = title_para.add_run("Technical Architecture Overview")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = "Calibri"
    title_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # Version line
    version_para = doc.add_paragraph()
    version_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    version_para.paragraph_format.space_after = Pt(2)
    version_run = version_para.add_run("Version 2.3")
    version_run.font.size = Pt(12)
    version_run.font.name = "Calibri"
    version_run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    # Date line
    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.space_after = Pt(12)
    date_run = date_para.add_run("March 15, 2026")
    date_run.font.size = Pt(11)
    date_run.font.name = "Calibri"
    date_run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    # ===== BODY CONTENT (no separator before it) =====

    # Section 1
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    h1_run = h1.add_run("1. System Overview")
    h1_run.bold = True
    h1_run.font.size = Pt(14)
    h1_run.font.name = "Calibri"
    h1_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(6)
    r1 = p1.add_run(
        "The platform is built on a microservices architecture deployed across "
        "three availability zones in AWS us-east-1. Each service communicates "
        "through an event-driven message bus (Apache Kafka) with guaranteed "
        "at-least-once delivery semantics. The API gateway handles approximately "
        "12,000 requests per second during peak hours."
    )
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(
        "Service discovery is managed through HashiCorp Consul, with health "
        "checks running at 10-second intervals. Circuit breakers are implemented "
        "using the Resilience4j library, configured with a failure rate threshold "
        "of 50% over a rolling window of 20 calls."
    )
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

    # Section 2
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("2. Data Layer")
    h2_run.bold = True
    h2_run.font.size = Pt(14)
    h2_run.font.name = "Calibri"
    h2_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    r3 = p3.add_run(
        "The primary datastore is PostgreSQL 15 with logical replication to "
        "read replicas. Write operations target the primary node while read "
        "queries are load-balanced across three replicas. Connection pooling "
        "is handled by PgBouncer with a maximum of 200 connections per pool."
    )
    r3.font.size = Pt(11)
    r3.font.name = "Calibri"

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(6)
    r4 = p4.add_run(
        "Redis 7.2 serves as the caching layer with a TTL policy of 300 seconds "
        "for session data and 3600 seconds for computed aggregates. Cache hit "
        "ratios consistently exceed 94% for the product catalog service."
    )
    r4.font.size = Pt(11)
    r4.font.name = "Calibri"

    # Section 3
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)
    h3_run = h3.add_run("3. Security and Compliance")
    h3_run.bold = True
    h3_run.font.size = Pt(14)
    h3_run.font.name = "Calibri"
    h3_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(6)
    r5 = p5.add_run(
        "All inter-service communication is encrypted using mutual TLS with "
        "certificates rotated every 90 days through the internal PKI managed "
        "by Vault. Authentication uses OAuth 2.0 with JWT tokens issued by "
        "the identity service, with token expiry set to 15 minutes."
    )
    r5.font.size = Pt(11)
    r5.font.name = "Calibri"

    p6 = doc.add_paragraph()
    p6.paragraph_format.space_after = Pt(6)
    r6 = p6.add_run(
        "Audit logging captures all state-changing operations and forwards "
        "events to a centralized SIEM (Splunk Enterprise). The system achieved "
        "SOC 2 Type II certification in January 2026 and undergoes quarterly "
        "penetration testing by CrowdStrike."
    )
    r6.font.size = Pt(11)
    r6.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
