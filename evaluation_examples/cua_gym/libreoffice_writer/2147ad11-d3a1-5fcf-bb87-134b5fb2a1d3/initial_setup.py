"""
Initial Setup: Create a Writer document with 15 pages of main content and an appendix
starting on page 16, all using Default Page Style with continuous numbering.
Task ID: writer_fs_081
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_081'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_footer_page_number(section):
    """Add a simple page number to the footer of a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # PAGE field code
    r1 = fp.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)

    r3 = fp.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Add footer with simple page numbers (continuous)
    add_footer_page_number(section)

    # --- Main Content (pages 1-15) ---
    # Title page
    title = doc.add_heading('Quarterly Business Performance Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sub.add_run('Prepared by the Strategic Planning Division')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Fiscal Year 2025 - Q3 Review')
    run.font.size = Pt(12)
    run.italic = True

    doc.add_page_break()

    # Page 2-3: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This report provides a comprehensive analysis of our company\'s performance '
        'during the third quarter of fiscal year 2025. The quarter was marked by '
        'significant growth in key market segments, particularly in the Asia-Pacific '
        'region where revenues increased by 23.4% year-over-year. Our strategic '
        'investments in digital transformation initiatives have begun to yield '
        'measurable returns, with operational efficiency improvements of 12.7% '
        'across our manufacturing divisions.'
    )
    doc.add_paragraph(
        'Key highlights include the successful launch of three new product lines, '
        'expansion into two previously untapped markets, and the completion of the '
        'CloudFirst migration project ahead of schedule and under budget. Customer '
        'satisfaction scores reached an all-time high of 94.2%, reflecting our '
        'commitment to service excellence and product quality. The total revenue '
        'for Q3 reached $287.5 million, exceeding projections by 8.3%.'
    )
    doc.add_paragraph(
        'However, challenges remain in the European market where regulatory changes '
        'have introduced additional compliance requirements. Supply chain disruptions '
        'in semiconductor components continue to affect our hardware division, though '
        'mitigation strategies implemented in Q2 have reduced the impact by approximately '
        '35% compared to the previous quarter.'
    )

    doc.add_page_break()

    # Pages 4-6: Financial Overview
    doc.add_heading('2. Financial Overview', level=1)
    doc.add_heading('2.1 Revenue Analysis', level=2)
    doc.add_paragraph(
        'Total consolidated revenue for Q3 2025 amounted to $287.5 million, representing '
        'a 15.2% increase from Q3 2024. This growth was primarily driven by our Software '
        'Solutions division, which contributed $142.3 million (49.5% of total revenue), '
        'followed by Professional Services at $89.7 million (31.2%), and Hardware Sales '
        'at $55.5 million (19.3%).'
    )

    # Revenue table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Division', 'Q3 2025 ($M)', 'Q3 2024 ($M)', 'YoY Growth']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data = [
        ['Software Solutions', '$142.3', '$118.9', '+19.7%'],
        ['Professional Services', '$89.7', '$78.4', '+14.4%'],
        ['Hardware Sales', '$55.5', '$52.1', '+6.5%'],
        ['Total Revenue', '$287.5', '$249.4', '+15.2%'],
        ['Operating Income', '$57.8', '$46.2', '+25.1%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_heading('2.2 Cost Structure', level=2)
    doc.add_paragraph(
        'Total operating expenses for the quarter were $229.7 million, resulting in an '
        'operating margin of 20.1%, an improvement of 1.6 percentage points from the '
        'prior year period. Cost of goods sold decreased by 3.2% as a percentage of '
        'revenue, reflecting improved supply chain management and favorable vendor '
        'renegotiations completed in early 2025.'
    )
    doc.add_paragraph(
        'Research and development expenditures totaled $38.4 million (13.4% of revenue), '
        'up from $31.2 million in Q3 2024. This increase reflects our continued commitment '
        'to innovation, with 47 active development projects across artificial intelligence, '
        'quantum computing readiness, and next-generation cybersecurity platforms.'
    )

    doc.add_page_break()

    doc.add_heading('2.3 Cash Flow and Balance Sheet', level=2)
    doc.add_paragraph(
        'Operating cash flow reached $62.3 million, a 28.4% increase over the comparable '
        'period. Free cash flow of $48.9 million provided ample liquidity for strategic '
        'investments and shareholder returns. The company ended the quarter with $312.7 '
        'million in cash and equivalents, and total debt of $145.0 million.'
    )
    doc.add_paragraph(
        'During the quarter, the company repurchased 1.2 million shares at an average price '
        'of $78.42, returning $94.1 million to shareholders through buybacks and dividends. '
        'The board approved an additional $200 million share repurchase authorization.'
    )

    doc.add_page_break()

    # Pages 7-9: Market Analysis
    doc.add_heading('3. Market Analysis', level=1)
    doc.add_heading('3.1 Regional Performance', level=2)
    doc.add_paragraph(
        'The Asia-Pacific region delivered exceptional results, with revenue growth of 23.4% '
        'driven by strong demand in Japan, South Korea, and emerging Southeast Asian markets. '
        'Our partnership with Tokyo Digital Solutions expanded our presence in the Japanese '
        'enterprise market, contributing $18.7 million in new contract value.'
    )
    doc.add_paragraph(
        'North America remained our largest market at $156.2 million (54.3% of total), '
        'growing at 11.8%. The federal sector showed particular strength with $34.5 million '
        'in new awards, including a $12.3 million multi-year contract with the Department '
        'of Energy for cybersecurity infrastructure modernization.'
    )
    doc.add_heading('3.2 Competitive Landscape', level=2)
    doc.add_paragraph(
        'Market share analysis indicates we maintained our leading position in the enterprise '
        'software segment with approximately 18.3% market share, up from 17.1% in Q3 2024. '
        'Key competitive wins include displacing incumbent vendors at three Fortune 500 '
        'companies, with total contract values exceeding $45 million over five-year terms.'
    )
    doc.add_paragraph(
        'The competitive environment intensified in the mid-market segment, where two new '
        'entrants launched competing products. Our response strategy includes enhanced '
        'product bundling, accelerated feature delivery, and a dedicated mid-market sales '
        'team established in August 2025.'
    )

    doc.add_page_break()

    doc.add_heading('3.3 Industry Trends', level=2)
    doc.add_paragraph(
        'Several macro trends continue to shape our market environment. The global digital '
        'transformation market is projected to reach $3.4 trillion by 2027, with a compound '
        'annual growth rate of 16.5%. Enterprise AI adoption accelerated significantly, with '
        '73% of enterprises now deploying AI in at least one business function, up from 58% '
        'a year ago.'
    )
    doc.add_paragraph(
        'Cybersecurity spending continues its upward trajectory, with global spending expected '
        'to exceed $215 billion in 2025. The increasing sophistication of threat actors and '
        'expanding regulatory requirements create substantial demand for our security solutions '
        'portfolio. Our Security Operations Center processed 2.3 billion events during Q3, '
        'identifying and mitigating 847 critical threats across our client base.'
    )

    doc.add_page_break()

    # Pages 10-12: Operations
    doc.add_heading('4. Operations Review', level=1)
    doc.add_heading('4.1 Product Development', level=2)
    doc.add_paragraph(
        'The product development organization achieved several significant milestones during Q3. '
        'The launch of CloudShield 5.0, our next-generation cybersecurity platform, received '
        'positive market reception with 234 new enterprise licenses in the first 45 days. '
        'DataSphere Analytics 3.2 introduced real-time streaming analytics capabilities, '
        'reducing time-to-insight by an average of 67% across beta customer deployments.'
    )
    doc.add_paragraph(
        'Our engineering teams delivered 1,847 feature enhancements and resolved 3,291 issues '
        'across all product lines. Sprint velocity improved by 14% following the adoption of '
        'our new agile framework, and automated test coverage reached 89.3%, up from 82.1% '
        'at the beginning of the fiscal year.'
    )
    doc.add_heading('4.2 Customer Success', level=2)
    doc.add_paragraph(
        'The Customer Success organization managed relationships with 2,847 active enterprise '
        'accounts during Q3. Net Revenue Retention Rate reached 118.7%, driven by strong '
        'expansion within existing accounts. Annual churn rate decreased to 4.2%, below our '
        'target of 5.0%, reflecting the effectiveness of our proactive engagement model.'
    )
    doc.add_paragraph(
        'Key account highlights include the expansion of our partnership with GlobalBank Corp, '
        'adding $8.3 million in annual recurring revenue through a comprehensive data platform '
        'deployment. Meridian Healthcare Systems extended their agreement for an additional '
        'five years with a 40% increase in scope, valued at $23.7 million.'
    )

    doc.add_page_break()

    doc.add_heading('4.3 Talent and Organization', level=2)
    doc.add_paragraph(
        'Total headcount reached 4,287 employees at quarter end, a net increase of 312 from '
        'Q2 2025. Key hires included Dr. Priya Sharma as Chief Data Officer, bringing 18 years '
        'of experience from leading technology firms, and James Rodriguez as VP of Asia-Pacific '
        'Sales, formerly with Oracle\'s enterprise division.'
    )
    doc.add_paragraph(
        'Employee engagement scores reached 82.4%, measured through our quarterly pulse survey. '
        'Voluntary turnover decreased to 8.7% annualized, well below the industry average of '
        '13.2%. Our diversity, equity, and inclusion initiatives resulted in a 15% increase in '
        'underrepresented group representation in leadership roles compared to the same period '
        'last year.'
    )

    doc.add_page_break()

    # Pages 13-15: Strategic Outlook
    doc.add_heading('5. Strategic Outlook', level=1)
    doc.add_heading('5.1 Q4 2025 Priorities', level=2)
    doc.add_paragraph(
        'As we enter the final quarter of fiscal year 2025, our strategic priorities are '
        'focused on three key areas: accelerating AI-powered product capabilities, expanding '
        'our global partner ecosystem, and optimizing operational efficiency to support '
        'sustained profitable growth.'
    )
    doc.add_paragraph('Key initiatives for Q4 include:')
    initiatives = [
        'Launch of AI Copilot integration across all major product lines by November 2025',
        'Expansion of the Global Partner Program with 50 new technology and consulting partners',
        'Completion of Phase 2 data center consolidation, targeting $12M in annual savings',
        'Release of CloudShield 5.1 with advanced threat intelligence and automated response',
        'Pilot deployment of quantum-safe encryption across government sector clients',
    ]
    for item in initiatives:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.2 Financial Guidance', level=2)
    doc.add_paragraph(
        'Based on current pipeline visibility and market conditions, we are raising our full-year '
        'revenue guidance to $1.12-1.14 billion, up from the previous range of $1.08-1.10 billion. '
        'Q4 revenue is expected to be in the range of $295-310 million, reflecting typical '
        'seasonal strength in enterprise purchasing cycles.'
    )
    doc.add_paragraph(
        'Operating margin is expected to remain in the 19-21% range for the full year, with '
        'potential upside from continued operational improvements and favorable product mix shift '
        'toward higher-margin software solutions. Capital expenditure is projected at $35-40 '
        'million for the full year, primarily for data center expansion and R&D infrastructure.'
    )

    doc.add_page_break()

    doc.add_heading('5.3 Long-Term Vision', level=2)
    doc.add_paragraph(
        'Our five-year strategic plan envisions the company as a leading provider of AI-powered '
        'enterprise solutions, with revenue exceeding $2.5 billion by fiscal year 2030. Key '
        'strategic pillars include organic product innovation, strategic acquisitions in adjacent '
        'markets, and geographic expansion into Latin America and the Middle East.'
    )
    doc.add_paragraph(
        'We believe the convergence of artificial intelligence, cloud computing, and cybersecurity '
        'creates a unique market opportunity that aligns perfectly with our capabilities and '
        'market position. Our investments in R&D, talent, and infrastructure position us to '
        'capitalize on these trends and deliver sustained value to our shareholders, customers, '
        'and employees.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'The management team remains confident in our ability to execute on this vision while '
        'maintaining disciplined financial management and a strong balance sheet. We look forward '
        'to updating our stakeholders on continued progress in our Q4 earnings report.'
    )

    # --- Appendix (starts on page 16) ---
    doc.add_page_break()

    doc.add_heading('Appendix', level=1)
    doc.add_heading('A. Detailed Financial Tables', level=2)
    doc.add_paragraph(
        'The following tables provide detailed breakdowns of financial data referenced '
        'throughout this report. All figures are in millions of US dollars unless '
        'otherwise noted.'
    )

    # Appendix table
    table2 = doc.add_table(rows=8, cols=5)
    table2.style = 'Table Grid'
    headers2 = ['Quarter', 'Revenue', 'COGS', 'Gross Profit', 'GP Margin']
    for i, h in enumerate(headers2):
        cell = table2.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    fin_data = [
        ['Q1 2024', '$231.8', '$138.2', '$93.6', '40.4%'],
        ['Q2 2024', '$240.5', '$141.9', '$98.6', '41.0%'],
        ['Q3 2024', '$249.4', '$145.7', '$103.7', '41.6%'],
        ['Q4 2024', '$268.3', '$153.4', '$114.9', '42.8%'],
        ['Q1 2025', '$259.2', '$149.8', '$109.4', '42.2%'],
        ['Q2 2025', '$271.8', '$155.2', '$116.6', '42.9%'],
        ['Q3 2025', '$287.5', '$162.3', '$125.2', '43.5%'],
    ]
    for r, row_data in enumerate(fin_data, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_heading('B. Regional Revenue Breakdown', level=2)

    table3 = doc.add_table(rows=6, cols=4)
    table3.style = 'Table Grid'
    headers3 = ['Region', 'Q3 2025 ($M)', 'Q3 2024 ($M)', 'Growth']
    for i, h in enumerate(headers3):
        cell = table3.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    reg_data = [
        ['North America', '$156.2', '$139.7', '+11.8%'],
        ['Europe', '$62.4', '$57.8', '+8.0%'],
        ['Asia-Pacific', '$48.3', '$39.1', '+23.4%'],
        ['Latin America', '$12.8', '$8.4', '+52.4%'],
        ['Middle East & Africa', '$7.8', '$4.4', '+77.3%'],
    ]
    for r, row_data in enumerate(reg_data, 1):
        for c, val in enumerate(row_data):
            table3.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_heading('C. Product Line Performance', level=2)
    doc.add_paragraph(
        'Detailed product performance metrics including customer acquisition costs, '
        'lifetime value ratios, and retention rates by product line are presented below. '
        'These metrics inform our investment prioritization framework and resource '
        'allocation decisions for the upcoming fiscal year planning cycle.'
    )
    doc.add_paragraph(
        'CloudShield platform demonstrated the strongest unit economics improvement, '
        'with customer acquisition cost decreasing by 18% while average contract value '
        'increased by 32%. DataSphere Analytics maintained the highest net revenue '
        'retention at 127%, reflecting strong cross-sell and upsell momentum within '
        'the installed base.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
