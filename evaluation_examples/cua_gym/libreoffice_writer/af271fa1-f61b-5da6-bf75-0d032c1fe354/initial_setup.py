"""
Initial Setup: Apply Internet Link character style to all URLs in a resource guide document
Task ID: osworld_writer_character_style_003
Domain: libreoffice_writer

Creates a resource guide document with 4 plain-text URLs (no character style applied).
The agent must apply the Internet Link character style to all 4 URLs.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_character_style_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading("Digital Resource Guide for Remote Teams", level=0)

    # Introduction paragraph
    intro = doc.add_paragraph(
        "This guide provides a curated collection of online resources for distributed teams "
        "working in technology, project management, and collaborative design. All resources "
        "listed below are freely accessible and regularly updated."
    )

    # Section 1: Project Management Tools
    doc.add_heading("1. Project Management & Collaboration", level=1)
    p1 = doc.add_paragraph(
        "For teams seeking a robust project management platform, Asana offers a comprehensive "
        "suite of task tracking and workflow automation features. You can get started by visiting "
        "their main portal at "
    )
    run_url1 = p1.add_run("https://www.asana.com/teams")
    # Deliberately NOT applying Internet Link style (initial state)
    p1.add_run(
        ". The platform supports Kanban boards, Gantt charts, and real-time team dashboards, "
        "making it ideal for cross-functional projects."
    )

    # Section 2: Version Control & DevOps
    doc.add_heading("2. Version Control & DevOps", level=1)
    p2 = doc.add_paragraph(
        "Modern software development relies heavily on distributed version control. GitHub "
        "provides both public and private repositories, along with CI/CD pipelines and "
        "integrated code review workflows. Teams can access the developer documentation at "
    )
    run_url2 = p2.add_run("https://docs.github.com/en/get-started")
    # Deliberately NOT applying Internet Link style (initial state)
    p2.add_run(
        ". This documentation covers branching strategies, pull request best practices, "
        "and GitHub Actions for automated deployments."
    )

    # Section 3: Design & Prototyping
    doc.add_heading("3. Design & Prototyping Resources", level=1)
    p3 = doc.add_paragraph(
        "UI/UX teams benefit greatly from collaborative design tools. Figma has become the "
        "industry standard for interface design and prototyping, offering real-time multiplayer "
        "editing. Their community resources and templates are available at "
    )
    run_url3 = p3.add_run("https://www.figma.com/community/resources")
    # Deliberately NOT applying Internet Link style (initial state)
    p3.add_run(
        ". Designers can browse thousands of free components, icon sets, and full UI kits "
        "contributed by the global design community."
    )

    # Section 4: Learning & Professional Development
    doc.add_heading("4. Learning & Professional Development", level=1)
    p4 = doc.add_paragraph(
        "Continuous learning is essential for technology professionals. Coursera partners with "
        "top universities and companies to offer accredited online courses and specializations. "
        "Browse the full catalog of technology and business courses at "
    )
    run_url4 = p4.add_run("https://www.coursera.org/browse/computer-science")
    # Deliberately NOT applying Internet Link style (initial state)
    p4.add_run(
        ". The platform offers both free auditing options and paid certificates that are "
        "recognized by leading employers worldwide."
    )

    # Closing section
    doc.add_heading("Usage Guidelines", level=1)
    closing = doc.add_paragraph(
        "All resources listed in this guide have been reviewed for quality and relevance. "
        "Links were verified as of Q1 2025. If you encounter broken links or wish to suggest "
        "additional resources, please contact the documentation team through the internal portal. "
        "This guide is updated on a quarterly basis to ensure accuracy and completeness."
    )

    doc.add_paragraph(
        "Note: Some external resources may require free registration before granting full access. "
        "Employees are encouraged to use their company email addresses when signing up for "
        "professional tools to consolidate licensing and support agreements."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
