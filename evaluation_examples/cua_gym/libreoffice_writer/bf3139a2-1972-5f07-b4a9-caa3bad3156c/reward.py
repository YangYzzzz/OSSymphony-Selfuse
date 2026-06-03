"""
Reward Script: Add SUM formula Total row at the bottom of a table
Task ID: writer_tbl_046
Domain: libreoffice_writer
Scoring:
  Component 1: New 'Total' row (row 6) exists at bottom of table (0.3 pts)
  Component 2: Cell B6 has SUM formula with cached value 700 (0.35 pts)
  Component 3: Cells C6 and D6 have SUM formulas with cached values 755 and 810 (0.35 pts)
Total: 1.0
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_046'

NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def cell_has_sum_formula(cell):
    """Return the formula text if a Word SUM field code is found in cell, else empty string."""
    for para in cell.paragraphs:
        instr_elems = para._element.findall('.//w:instrText', NS)
        for elem in instr_elems:
            if elem.text and 'SUM' in elem.text.upper():
                return elem.text.strip()
    return ''


def verify_task(file_path):
    """
    Verify that a new Total row with SUM formulas has been added to the bottom
    of the table in the quarterly_sales document.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: at least 1 table must exist
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)

    # Component 1: New 'Total' row exists as row 6 (index 5) in the table (0.3 points)
    # Initial env has 5 rows; golden env must have 6 rows with 'Total' in column A of last row.
    try:
        if num_rows >= 6:
            last_row = table.rows[5]
            cell_a6_text = last_row.cells[0].text.strip()
            if cell_a6_text.lower() == 'total':
                print(f"PASS: Component 1 — 'Total' row found at row 6, cell A6='{cell_a6_text}' (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 1 — Row 6 exists but cell A6='{cell_a6_text}', expected 'Total'")
        else:
            print(f"FAIL: Component 1 — Table has {num_rows} rows, expected 6 (no Total row added)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Cell B6 contains a SUM field code and cached value is 700 (0.35 points)
    # We check for the field code instrText containing 'SUM' and the cell text showing '700'.
    try:
        if num_rows >= 6:
            last_row = table.rows[5]
            cell_b6 = last_row.cells[1]
            cell_b6_text = cell_b6.text.strip()
            b6_formula = cell_has_sum_formula(cell_b6)

            if b6_formula and cell_b6_text == '700':
                print(f"PASS: Component 2 — B6 has SUM formula ('{b6_formula}') and value '700' (0.35 pts)")
                total_score += 0.35
            elif b6_formula and cell_b6_text != '700':
                print(f"FAIL: Component 2 — B6 has SUM formula but cached value='{cell_b6_text}', expected '700'")
            elif not b6_formula and cell_b6_text == '700':
                print(f"FAIL: Component 2 — B6 has correct value '700' but no SUM field code detected")
            else:
                print(f"FAIL: Component 2 — B6: no SUM formula, value='{cell_b6_text}' (expected SUM formula + '700')")
        else:
            print(f"FAIL: Component 2 — Table has {num_rows} rows, no row 6 to check B6")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Cells C6 and D6 contain SUM field codes with values 755 and 810 (0.35 points)
    # Both cells must have SUM formulas and correct cached values.
    try:
        if num_rows >= 6:
            last_row = table.rows[5]
            cell_c6 = last_row.cells[2]
            cell_c6_text = cell_c6.text.strip()
            c6_formula = cell_has_sum_formula(cell_c6)

            cell_d6 = last_row.cells[3]
            cell_d6_text = cell_d6.text.strip()
            d6_formula = cell_has_sum_formula(cell_d6)

            c6_correct = (len(c6_formula) > 0) and (cell_c6_text == '755')
            d6_correct = (len(d6_formula) > 0) and (cell_d6_text == '810')

            if c6_correct and d6_correct:
                print(f"PASS: Component 3 — C6 has SUM formula with value '755', D6 has SUM formula with value '810' (0.35 pts)")
                total_score += 0.35
            else:
                if not c6_correct:
                    print(f"FAIL: Component 3 — C6: formula='{c6_formula}', value='{cell_c6_text}' (expected SUM formula + '755')")
                if not d6_correct:
                    print(f"FAIL: Component 3 — D6: formula='{d6_formula}', value='{cell_d6_text}' (expected SUM formula + '810')")
        else:
            print(f"FAIL: Component 3 — Table has {num_rows} rows, no row 6 to check C6/D6")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
