"""
Initial Setup: Quarterly Sales Document with Table
Task ID: writer_tbl_046
Domain: libreoffice_writer

Creates a .docx file with a 5-row, 4-column table of quarterly sales data.
No Total row exists yet — the agent must add it.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_046'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a document title
    title = doc.add_heading('Quarterly Sales Report', level=1)

    # Add a brief intro paragraph
    doc.add_paragraph('The following table summarizes product sales performance across Q1, Q2, and Q3.')

    # Create the sales table: 5 rows (1 header + 4 data), 4 columns
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    # Header row
    header_data = ['Product', 'Q1', 'Q2', 'Q3']
    header_row = table.rows[0]
    for col_idx, header_text in enumerate(header_data):
        cell = header_row.cells[col_idx]
        cell.text = header_text
        # Make header text bold
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows (no Total row — agent must add it)
    data_rows = [
        ['Laptops',     '120', '150', '135'],
        ['Tablets',     '80',  '95',  '110'],
        ['Phones',      '200', '230', '215'],
        ['Accessories', '300', '280', '350'],
    ]

    for row_idx, row_data in enumerate(data_rows, start=1):
        row = table.rows[row_idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
