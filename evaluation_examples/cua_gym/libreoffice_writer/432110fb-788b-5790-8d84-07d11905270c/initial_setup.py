"""
Initial Setup: Change text wrapping of diagram image from None to Parallel
Task ID: writer_frd_069
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_069'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_diagram_image():
    """Create a simple technical diagram image using Pillow."""
    from PIL import Image, ImageDraw, ImageFont

    width, height = 480, 320
    img = Image.new('RGB', (width, height), color='#F0F4F8')
    draw = ImageDraw.Draw(img)

    # Draw a simple system architecture diagram
    # Box 1: Client
    draw.rectangle([40, 40, 180, 100], outline='#2C5F8A', width=2, fill='#D6E8F7')
    draw.text((75, 60), "Client App", fill='#1A3A5C')

    # Box 2: API Gateway
    draw.rectangle([200, 40, 340, 100], outline='#2C5F8A', width=2, fill='#D6E8F7')
    draw.text((220, 60), "API Gateway", fill='#1A3A5C')

    # Box 3: Backend
    draw.rectangle([200, 140, 340, 200], outline='#2C5F8A', width=2, fill='#D6E8F7')
    draw.text((225, 160), "Backend", fill='#1A3A5C')

    # Box 4: Database
    draw.rectangle([200, 240, 340, 300], outline='#2C5F8A', width=2, fill='#D6E8F7')
    draw.text((225, 260), "Database", fill='#1A3A5C')

    # Arrows
    draw.line([180, 70, 200, 70], fill='#2C5F8A', width=2)
    draw.line([270, 100, 270, 140], fill='#2C5F8A', width=2)
    draw.line([270, 200, 270, 240], fill='#2C5F8A', width=2)

    # Arrow heads
    draw.polygon([(200, 65), (195, 70), (200, 75)], fill='#2C5F8A')
    draw.polygon([(265, 140), (270, 135), (275, 140)], fill='#2C5F8A')
    draw.polygon([(265, 240), (270, 235), (275, 240)], fill='#2C5F8A')

    # Title
    draw.text((140, 5), "System Architecture", fill='#1A3A5C')

    img_path = f'{WORKDIR}/diagram.png'
    img.save(img_path)
    return img_path


def create_initial():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    # Create the diagram image first
    img_path = create_diagram_image()

    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading('Technical Architecture Document', level=1)

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=2)
    p1 = doc.add_paragraph()
    p1.add_run(
        'This document describes the high-level architecture of the Meridian '
        'Analytics Platform, a cloud-native data processing system designed to '
        'handle real-time telemetry streams from distributed IoT sensors across '
        'manufacturing facilities in North America and Europe.'
    )

    p2 = doc.add_paragraph()
    p2.add_run(
        'The platform was initially deployed in Q3 2024 at the Stuttgart facility '
        'and has since been expanded to cover 14 production lines across 5 locations. '
        'Current throughput exceeds 2.4 million events per second during peak '
        'operational windows.'
    )

    # --- Section 2: System Overview ---
    doc.add_heading('2. System Overview', level=2)
    p3 = doc.add_paragraph()
    p3.add_run(
        'The following diagram illustrates the core components and their '
        'interactions within the Meridian Analytics Platform. Each service '
        'communicates through a message broker to ensure loose coupling and '
        'horizontal scalability.'
    )

    # --- Insert diagram image INLINE (no wrapping) ---
    # add_picture inserts as inline by default (no text wrap)
    doc.add_picture(img_path, width=Inches(4.5))
    # Center the image paragraph
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_cap = p_caption.add_run('Figure 1: System Architecture Overview')
    run_cap.italic = True
    run_cap.font.size = Pt(9)
    run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Section 3: Component Details ---
    doc.add_heading('3. Component Details', level=2)

    doc.add_heading('3.1 Client Application Layer', level=3)
    p4 = doc.add_paragraph()
    p4.add_run(
        'The client application is built using React 18 with TypeScript, providing '
        'real-time dashboard visualizations for facility managers. WebSocket '
        'connections maintain persistent communication with the API gateway, '
        'allowing sub-second updates to monitoring panels.'
    )

    doc.add_heading('3.2 API Gateway', level=3)
    p5 = doc.add_paragraph()
    p5.add_run(
        'All external traffic routes through a Kong-based API gateway deployed on '
        'Kubernetes. The gateway handles authentication via OAuth 2.0 tokens, rate '
        'limiting at 10,000 requests per minute per client, and request routing to '
        'appropriate backend microservices.'
    )

    doc.add_heading('3.3 Backend Services', level=3)
    p6 = doc.add_paragraph()
    p6.add_run(
        'The backend consists of 12 microservices written in Go and Python. The '
        'ingestion service processes raw sensor data at approximately 180,000 events '
        'per second per instance. The transformation pipeline applies calibration '
        'curves and anomaly detection models before persisting results.'
    )

    doc.add_heading('3.4 Data Storage', level=3)
    p7 = doc.add_paragraph()
    p7.add_run(
        'Time-series data is stored in TimescaleDB with a 90-day hot retention '
        'window and automatic tiering to S3-compatible object storage for long-term '
        'archival. The current dataset spans approximately 4.2 TB of compressed '
        'telemetry records dating back to the initial deployment.'
    )

    # --- Section 4: Deployment ---
    doc.add_heading('4. Deployment Architecture', level=2)
    p8 = doc.add_paragraph()
    p8.add_run(
        'Production workloads run on a dedicated Kubernetes cluster with 24 worker '
        'nodes (each 32 vCPU, 128 GB RAM) across three availability zones. CI/CD '
        'pipelines are managed through GitLab with automated canary deployments '
        'requiring approval gates for production rollouts.'
    )

    p9 = doc.add_paragraph()
    p9.add_run(
        'Monitoring is provided by a Prometheus/Grafana stack with custom alerting '
        'rules for SLA compliance. Mean time to detection for critical incidents is '
        'currently 47 seconds, with automated remediation covering 78% of known '
        'failure modes.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Clean up temp image
    os.remove(img_path)

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
