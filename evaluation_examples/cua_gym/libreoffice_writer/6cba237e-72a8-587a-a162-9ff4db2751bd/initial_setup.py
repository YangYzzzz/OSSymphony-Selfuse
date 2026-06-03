"""
Initial Setup: Offer letter with compensation details as plain paragraph text
Task ID: writer_hr_052
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_052'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Company letterhead ---
    heading = doc.add_heading('Meridian Technologies Inc.', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('1200 Innovation Drive, Suite 400\nSan Francisco, CA 94107')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Date and recipient ---
    doc.add_paragraph('')  # spacer
    date_para = doc.add_paragraph('March 18, 2025')
    date_para.paragraph_format.space_after = Pt(6)

    doc.add_paragraph('Ms. Elena Rodriguez')
    doc.add_paragraph('742 Westbrook Lane')
    last_addr = doc.add_paragraph('Portland, OR 97201')
    last_addr.paragraph_format.space_after = Pt(12)

    # --- Greeting ---
    greeting = doc.add_paragraph('Dear Ms. Rodriguez,')
    greeting.paragraph_format.space_after = Pt(12)

    # --- Opening paragraph ---
    opening = doc.add_paragraph(
        'We are pleased to extend this offer of employment for the position of '
        'Senior Software Engineer at Meridian Technologies Inc. We were impressed '
        'by your technical expertise and collaborative approach during the interview '
        'process, and we believe you will be an excellent addition to our engineering team.'
    )
    opening.paragraph_format.space_after = Pt(6)

    # --- Role details ---
    role_para = doc.add_paragraph(
        'You will be reporting to David Park, VP of Engineering, and your start date '
        'is scheduled for April 14, 2025. Your primary office location will be our '
        'San Francisco headquarters, with the option to work remotely up to two days '
        'per week.'
    )
    role_para.paragraph_format.space_after = Pt(12)

    # --- Compensation section heading ---
    comp_heading = doc.add_heading('Compensation and Benefits', level=2)
    for run in comp_heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)

    # --- Compensation details as PLAIN TEXT (NOT a table) ---
    comp_intro = doc.add_paragraph(
        'We are offering the following compensation package as part of your employment:'
    )
    comp_intro.paragraph_format.space_after = Pt(6)

    comp_items = [
        'Base Salary: $125,000 per annum, paid on a semi-monthly basis.',
        'Annual Bonus Target: 15% of your base salary, subject to individual and company performance metrics.',
        'Stock Options: 5,000 shares vesting over a 4-year period with a one-year cliff, subject to the terms of the company equity plan.',
        'Health Insurance: Comprehensive medical, dental, and vision coverage. The company pays 80% of premiums for employee and dependents.',
        '401(k) Match: The company matches employee contributions up to 6% of your annual salary.',
        'Paid Time Off (PTO): 20 days of PTO per year, accrued on a per-pay-period basis, in addition to company-observed holidays.'
    ]

    for item in comp_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)

    # --- Additional terms ---
    doc.add_paragraph('')  # spacer

    addl_heading = doc.add_heading('Additional Terms', level=2)
    for run in addl_heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)

    addl_text = doc.add_paragraph(
        'This offer is contingent upon successful completion of a background check '
        'and verification of your eligibility to work in the United States. Employment '
        'at Meridian Technologies is at-will, meaning either party may terminate the '
        'employment relationship at any time, with or without cause or notice.'
    )
    addl_text.paragraph_format.space_after = Pt(6)

    nda_text = doc.add_paragraph(
        'On your first day, you will be asked to sign a Confidentiality and '
        'Intellectual Property Agreement, as well as acknowledge receipt of the '
        'Employee Handbook.'
    )
    nda_text.paragraph_format.space_after = Pt(12)

    # --- Closing ---
    closing = doc.add_paragraph(
        'We are excited to welcome you to Meridian Technologies. Please confirm your '
        'acceptance by signing and returning this letter by March 28, 2025. Should you '
        'have any questions, feel free to contact me directly at (415) 555-0192 or '
        'hr@meridiantech.com.'
    )
    closing.paragraph_format.space_after = Pt(18)

    doc.add_paragraph('Sincerely,')
    doc.add_paragraph('')
    doc.add_paragraph('Jennifer Walsh')
    doc.add_paragraph('Director of Human Resources')
    doc.add_paragraph('Meridian Technologies Inc.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
