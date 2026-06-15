"""
Initial Setup: Create unstructured report document with all content but no formatting
Task ID: writer_biz_058
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_058'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # All content as plain Normal paragraphs, no heading styles, no page breaks
    # Title page content (just plain text, no formatting)
    doc.add_paragraph("Quarterly Business Performance Report")
    doc.add_paragraph("Prepared by: Strategic Analytics Division")
    doc.add_paragraph("Meridian Global Consulting Group")
    doc.add_paragraph("Report Date: March 2025")
    doc.add_paragraph("Confidential — For Internal Distribution Only")
    doc.add_paragraph("")

    # TOC placeholder (just plain text)
    doc.add_paragraph("Table of Contents")
    doc.add_paragraph("[Table of Contents will be generated here]")
    doc.add_paragraph("")

    # Executive Summary
    doc.add_paragraph("Executive Summary")
    doc.add_paragraph(
        "This report presents a comprehensive analysis of Meridian Global Consulting Group's "
        "performance during Q4 2024 and provides strategic recommendations for Q1 2025. "
        "Overall revenue grew by 12.3% year-over-year, driven primarily by expansion in the "
        "Asia-Pacific region and strong demand for digital transformation advisory services."
    )
    doc.add_paragraph(
        "Key highlights include a 15% increase in client retention rates, successful onboarding "
        "of 47 new enterprise accounts, and the launch of three new service verticals. However, "
        "operating margins declined slightly due to increased hiring costs and infrastructure "
        "investments in our cloud analytics platform."
    )
    doc.add_paragraph(
        "The executive team recommends prioritizing operational efficiency improvements, "
        "targeted expansion into the European financial services sector, and continued "
        "investment in AI-powered consulting tools for the upcoming quarter."
    )
    doc.add_paragraph("")

    # Body Section 1
    doc.add_paragraph("Financial Performance Analysis")
    doc.add_paragraph(
        "Total revenue for Q4 2024 reached $48.7 million, representing a 12.3% increase "
        "compared to the same period in the prior year. The consulting services segment "
        "contributed $31.2 million, while technology implementation services accounted for "
        "$17.5 million. This shift toward technology services reflects the broader market "
        "trend of integrated consulting and technology delivery."
    )
    doc.add_paragraph(
        "Regional performance varied significantly. North America generated $22.1 million "
        "(45.4% of total), Asia-Pacific contributed $15.8 million (32.4%), and EMEA accounted "
        "for $10.8 million (22.2%). The Asia-Pacific region showed the strongest growth at "
        "18.7%, largely attributed to new engagements with major financial institutions in "
        "Singapore, Tokyo, and Sydney."
    )
    doc.add_paragraph(
        "Operating expenses totaled $39.4 million, yielding an operating margin of 19.1%, "
        "down from 21.3% in Q4 2023. The decline is primarily attributable to a $2.8 million "
        "increase in personnel costs associated with hiring 112 new consultants and a $1.4 "
        "million investment in the Meridian Analytics Cloud platform."
    )
    doc.add_paragraph("")

    # Body Section 2
    doc.add_paragraph("Client Engagement and Market Position")
    doc.add_paragraph(
        "Client satisfaction scores remained strong at 4.6 out of 5.0, consistent with our "
        "industry-leading performance over the past six quarters. The Net Promoter Score "
        "improved to 72, up from 68 in the previous quarter, reflecting positive reception of "
        "our enhanced delivery methodology and dedicated client success teams."
    )
    doc.add_paragraph(
        "During Q4, we successfully onboarded 47 new enterprise accounts, including notable "
        "wins with TransGlobal Insurance, Pacific Rim Holdings, and the European Central "
        "Banking Consortium. Our average contract value for new engagements increased to "
        "$1.2 million, up 8% from Q3 2024, indicating successful movement upmarket."
    )
    doc.add_paragraph(
        "Competitive analysis indicates that Meridian continues to hold the third-largest "
        "market share in the mid-market consulting segment at 8.2%. Our primary competitors, "
        "Nexus Advisory Group (11.4%) and Pinnacle Strategy Partners (9.7%), have both "
        "signaled expansion into our strongest verticals, necessitating proactive defensive "
        "positioning."
    )
    doc.add_paragraph("")

    # Body Section 3
    doc.add_paragraph("Strategic Recommendations")
    doc.add_paragraph(
        "Based on the analysis presented in this report, the Strategic Analytics Division "
        "proposes the following initiatives for Q1 2025:"
    )
    doc.add_paragraph(
        "First, establish a dedicated European Financial Services practice led by Senior "
        "Partner Rachel Thornton, targeting $5 million in new revenue by Q3 2025. Initial "
        "focus should be on regulatory compliance consulting for MiFID III implementation, "
        "an area where current market coverage is fragmented."
    )
    doc.add_paragraph(
        "Second, accelerate the Meridian Analytics Cloud platform rollout by allocating an "
        "additional $2 million in development resources. Early adopter feedback indicates "
        "potential for 30% improvement in consultant productivity, which would directly "
        "address the operating margin compression observed this quarter."
    )
    doc.add_paragraph(
        "Third, implement a structured talent development program to reduce consultant "
        "attrition from the current 14% to below 10% within 12 months. Estimated cost "
        "savings from improved retention: $3.6 million annually."
    )
    doc.add_paragraph("")

    # Appendix
    doc.add_paragraph("Appendix")
    doc.add_paragraph(
        "A. Detailed Financial Statements — Complete income statement, balance sheet, and "
        "cash flow analysis for Q4 2024 are available upon request from the Finance Division."
    )
    doc.add_paragraph(
        "B. Client Survey Methodology — Satisfaction scores were collected via anonymous "
        "online survey administered to 834 client contacts across all active engagements. "
        "Response rate: 62.4%."
    )
    doc.add_paragraph(
        "C. Competitive Landscape Data — Market share estimates derived from Gartner Q4 2024 "
        "Market Share Analysis for Management Consulting Services (Published January 2025)."
    )
    doc.add_paragraph(
        "D. Regional Performance Breakdown — Supplementary tables with country-level revenue "
        "data, headcount, and utilization rates are maintained in the Strategic Analytics "
        "shared repository."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
