"""
Reward Script: Professional report template with heading styles, page breaks, and title page
Task ID: writer_biz_058
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Heading 1 styles applied to section headings
  Component 2 (0.30): Page break before on body section headings (pbb=True)
  Component 3 (0.20): Title page formatting (bold, large font, centered)
  Component 4 (0.15): Page breaks separating title page / TOC / executive summary
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_058'

# Known section heading texts that should be Heading 1 in the golden doc
EXPECTED_HEADINGS = [
    'Table of Contents',
    'Executive Summary',
    'Financial Performance Analysis',
    'Client Engagement and Market Position',
    'Strategic Recommendations',
    'Appendix',
]

# Headings that must have page_break_before = True (body sections + appendix)
EXPECTED_PBB_HEADINGS = [
    'Financial Performance Analysis',
    'Client Engagement and Market Position',
    'Strategic Recommendations',
    'Appendix',
]


def persist_app_state(domain):
    """Attempt to save any open LibreOffice document."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        print(f"CRITICAL: Cannot import python-docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Build lookup: heading text -> (style_name, page_break_before)
    heading_info = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in EXPECTED_HEADINGS:
            style_name = para.style.name if para.style else 'None'
            pbb = para.paragraph_format.page_break_before
            heading_info[text] = (style_name, pbb)

    # =========================================================
    # Component 1: Heading 1 styles applied to section headings (0.35 pts)
    # At least 5 of the 6 expected headings should use Heading 1
    # =========================================================
    try:
        heading1_count = 0
        for htext in EXPECTED_HEADINGS:
            if htext in heading_info:
                style_name, _ = heading_info[htext]
                if style_name == 'Heading 1':
                    heading1_count += 1
                    print(f"  FOUND Heading 1: '{htext}'")
                else:
                    print(f"  MISS: '{htext}' has style '{style_name}', expected 'Heading 1'")
            else:
                print(f"  MISS: '{htext}' not found in document")

        if heading1_count >= 5:
            print(f"PASS: Component 1 — {heading1_count}/6 headings use Heading 1 (0.35 pts)")
            total_score += 0.35
        elif heading1_count >= 3:
            partial = round(0.35 * (heading1_count / 6), 2)
            print(f"PARTIAL: Component 1 — {heading1_count}/6 headings use Heading 1 ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — only {heading1_count}/6 headings use Heading 1")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # =========================================================
    # Component 2: Page break before on body section headings (0.30 pts)
    # The 3 body sections + Appendix should have page_break_before=True
    # =========================================================
    try:
        pbb_count = 0
        for htext in EXPECTED_PBB_HEADINGS:
            if htext in heading_info:
                _, pbb = heading_info[htext]
                if pbb:
                    pbb_count += 1
                    print(f"  FOUND pbb=True: '{htext}'")
                else:
                    print(f"  MISS pbb: '{htext}' has page_break_before={pbb}")
            else:
                print(f"  MISS pbb: '{htext}' not found")

        if pbb_count >= 4:
            print(f"PASS: Component 2 — {pbb_count}/4 headings have page_break_before (0.30 pts)")
            total_score += 0.30
        elif pbb_count >= 2:
            partial = round(0.30 * (pbb_count / 4), 2)
            print(f"PARTIAL: Component 2 — {pbb_count}/4 headings have page_break_before ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — only {pbb_count}/4 headings have page_break_before")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =========================================================
    # Component 3: Title page formatting (0.20 pts)
    # Title text should be bold, large (>=20pt), and centered
    # =========================================================
    try:
        title_formatted = False
        title_text = 'Quarterly Business Performance Report'
        for para in doc.paragraphs:
            if title_text in para.text:
                # Check alignment
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                alignment = para.paragraph_format.alignment
                is_centered = alignment == WD_PARAGRAPH_ALIGNMENT.CENTER

                # Check run formatting
                is_bold = False
                is_large = False
                for run in para.runs:
                    if run.font.bold:
                        is_bold = True
                    if run.font.size and run.font.size.pt >= 20:
                        is_large = True

                checks_passed = sum([is_centered, is_bold, is_large])
                print(f"  Title: centered={is_centered}, bold={is_bold}, large_font={is_large}")

                if checks_passed == 3:
                    title_formatted = True
                    print(f"PASS: Component 3 — Title fully formatted (0.20 pts)")
                    total_score += 0.20
                elif checks_passed >= 1:
                    partial = round(0.20 * (checks_passed / 3), 2)
                    print(f"PARTIAL: Component 3 — Title {checks_passed}/3 format checks ({partial} pts)")
                    total_score += partial
                    title_formatted = True
                break

        if not title_formatted:
            print(f"FAIL: Component 3 — Title paragraph not found or no formatting applied")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # =========================================================
    # Component 4: Page breaks separating title/TOC/exec summary sections (0.15 pts)
    # There should be run-level page breaks OR section breaks between
    # title page -> TOC and TOC -> Executive Summary
    # =========================================================
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # Count run-level page breaks in the document
        run_page_breaks = 0
        for para in doc.paragraphs:
            for run in para.runs:
                for br in run.element.findall('.//w:br', ns):
                    btype = br.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'line')
                    if btype == 'page':
                        run_page_breaks += 1

        # Also count section breaks beyond the first section
        extra_sections = max(0, len(doc.sections) - 1)

        # We also count page_break_before on TOC and Executive Summary headings
        toc_pbb = False
        exec_pbb = False
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt == 'Table of Contents' and para.paragraph_format.page_break_before:
                toc_pbb = True
            if txt == 'Executive Summary' and para.paragraph_format.page_break_before:
                exec_pbb = True

        # Total page separation indicators (run breaks + section breaks + pbb on TOC/Exec)
        separation_count = run_page_breaks + extra_sections + int(toc_pbb) + int(exec_pbb)

        print(f"  Run page breaks: {run_page_breaks}, extra sections: {extra_sections}, TOC pbb: {toc_pbb}, Exec pbb: {exec_pbb}")

        if separation_count >= 2:
            print(f"PASS: Component 4 — {separation_count} page separations found (0.15 pts)")
            total_score += 0.15
        elif separation_count >= 1:
            print(f"PARTIAL: Component 4 — {separation_count} page separation found (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 4 — no page separations between title/TOC/exec summary")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
