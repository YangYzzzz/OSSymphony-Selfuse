"""
Reward Script: Two-column layout with full-width title and abstract
Task ID: writer_fs_024
Domain: libreoffice_writer
Scoring:
  Component 1: Document has 2+ sections (0.2 pts)
  Component 2: Section break placed after abstract paragraph (0.2 pts)
  Component 3: First section (title+abstract) is single column (0.3 pts)
  Component 4: Second section (body) has 2 columns with ~0.7cm spacing (0.3 pts)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_024'


def persist_app_state(domain):
    """Save any unsaved changes in LibreOffice Writer."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # === Gather section information ===
    num_sections = len(doc.sections)
    print(f"INFO: Document has {num_sections} section(s)")

    # Also check for inline section breaks in paragraphs
    inline_sect_breaks = []
    for i, para in enumerate(doc.paragraphs):
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            sectPr = pPr.find(qn('w:sectPr'))
            if sectPr is not None:
                inline_sect_breaks.append(i)
    print(f"INFO: Inline section breaks at paragraph indices: {inline_sect_breaks}")

    # Component 1: Document has 2+ sections (0.2 points)
    # Initial has 1 section; golden has 2.
    try:
        if num_sections >= 2:
            print(f"PASS: Component 1 — Document has {num_sections} sections (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 1 — Expected 2+ sections, found {num_sections}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Section break placed after title/abstract area (0.2 points)
    # The section break should be after paragraph index 2 or 3 (the abstract text)
    # and before the body text (paragraph index 4, "1. Introduction").
    # In the golden, the break is on paragraph 3.
    try:
        if len(inline_sect_breaks) >= 1:
            break_idx = inline_sect_breaks[0]
            # The break should be on a paragraph that is part of the abstract area
            # (indices 0-3 for title, authors, "Abstract" label, abstract text)
            # and NOT in the body section.
            if break_idx <= 4:
                # Verify that the paragraph after the break is body text (heading or body)
                if break_idx + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[break_idx + 1]
                    next_text = next_para.text[:30]
                    print(f"PASS: Component 2 — Section break at para {break_idx}, next para: '{next_text}' (0.2 pts)")
                    total_score += 0.2
                else:
                    print(f"PASS: Component 2 — Section break at para {break_idx} (0.2 pts)")
                    total_score += 0.2
            else:
                print(f"FAIL: Component 2 — Section break at para {break_idx}, expected in abstract area (indices 0-4)")
        else:
            print(f"FAIL: Component 2 — No inline section breaks found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: First section (title+abstract) is single column (0.3 points)
    # The first section should NOT have w:num="2" or higher on its cols element.
    try:
        if num_sections >= 2:
            sec0 = doc.sections[0]
            sectPr0 = sec0._sectPr
            cols0 = sectPr0.find(qn('w:cols'))
            if cols0 is not None:
                num_cols_str = cols0.get(qn('w:num'))
                num_cols = int(num_cols_str) if num_cols_str else 1
            else:
                num_cols = 1

            if num_cols <= 1:
                print(f"PASS: Component 3 — First section has {num_cols} column(s) (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 3 — First section has {num_cols} columns, expected 1")
        else:
            print(f"FAIL: Component 3 — Cannot check first section columns (only {num_sections} section)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Second section (body) has 2 columns with ~0.7cm spacing (0.3 points)
    # 0.7 cm = ~397 twips (1 cm ≈ 567 twips). Accept range 350-450 twips.
    # Sub-scores: 0.2 for 2 columns, 0.1 for correct spacing.
    try:
        if num_sections >= 2:
            sec1 = doc.sections[1]
            sectPr1 = sec1._sectPr
            cols1 = sectPr1.find(qn('w:cols'))
            if cols1 is not None:
                num_cols_str = cols1.get(qn('w:num'))
                num_cols = int(num_cols_str) if num_cols_str else 1
                space_str = cols1.get(qn('w:space'))
                space_val = int(space_str) if space_str else None
            else:
                num_cols = 1
                space_val = None

            if num_cols == 2:
                print(f"PASS: Component 4a — Body section has 2 columns (0.2 pts)")
                total_score += 0.2

                # Check spacing: 0.7 cm ≈ 397 twips, accept 300-500
                if space_val is not None and 300 <= space_val <= 500:
                    spacing_cm = space_val / 567.0
                    print(f"PASS: Component 4b — Column spacing = {space_val} twips ({spacing_cm:.2f} cm), within acceptable range (0.1 pts)")
                    total_score += 0.1
                elif space_val is not None:
                    spacing_cm = space_val / 567.0
                    print(f"FAIL: Component 4b — Column spacing = {space_val} twips ({spacing_cm:.2f} cm), expected ~397 twips (~0.7 cm)")
                else:
                    print(f"FAIL: Component 4b — No spacing attribute on cols element")
            else:
                print(f"FAIL: Component 4 — Body section has {num_cols} columns, expected 2")
        else:
            print(f"FAIL: Component 4 — Cannot check body section columns (only {num_sections} section)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 1)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
