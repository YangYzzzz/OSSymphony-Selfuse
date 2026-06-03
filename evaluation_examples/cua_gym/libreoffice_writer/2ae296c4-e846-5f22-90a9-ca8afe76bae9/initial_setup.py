"""
Initial Setup: Create a single-column academic paper document with title, abstract, and body text.
Task ID: writer_fs_024
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_024'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title ---
    title_para = doc.add_heading(
        "Investigating the Effects of Urban Green Spaces on Mental Health Outcomes: "
        "A Longitudinal Study",
        level=0
    )
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(6)

    # --- Authors ---
    authors = doc.add_paragraph()
    authors.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    authors.paragraph_format.space_after = Pt(12)
    run = authors.add_run("Elena Vasquez, David Kim, and Rachel Okonkwo")
    run.font.size = Pt(11)
    run.italic = True

    # --- Abstract heading ---
    abs_heading = doc.add_paragraph()
    abs_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    abs_heading.paragraph_format.space_before = Pt(12)
    abs_heading.paragraph_format.space_after = Pt(6)
    run_abs = abs_heading.add_run("Abstract")
    run_abs.bold = True
    run_abs.font.size = Pt(12)

    # --- Abstract text ---
    abstract = doc.add_paragraph()
    abstract.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    abstract.paragraph_format.space_after = Pt(12)
    run_abs_text = abstract.add_run(
        "This study examines the relationship between proximity to urban green spaces "
        "and mental health outcomes among 2,847 residents across twelve metropolitan areas "
        "over a five-year period (2019\u20132024). Using a mixed-methods longitudinal design, "
        "we assessed self-reported well-being scores, clinical anxiety and depression "
        "indicators, and physiological stress markers at six-month intervals. Results "
        "indicate that participants living within 500 meters of maintained green spaces "
        "showed a statistically significant reduction in generalized anxiety disorder "
        "symptoms (p < 0.001) and a 23% improvement in subjective well-being scores "
        "compared to control groups. The effect was particularly pronounced among "
        "participants aged 25\u201345 in high-density residential zones. These findings "
        "suggest that urban planning policies prioritizing accessible green infrastructure "
        "may serve as a cost-effective public health intervention."
    )
    run_abs_text.font.size = Pt(11)

    # --- Body text: Introduction ---
    intro_heading = doc.add_heading("1. Introduction", level=1)
    intro_heading.paragraph_format.space_before = Pt(18)

    intro_p1 = doc.add_paragraph()
    intro_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = intro_p1.add_run(
        "The rapid urbanization of the twenty-first century has fundamentally altered "
        "the built environments in which the majority of the global population resides. "
        "By 2023, approximately 56% of the world's population lived in urban areas, a "
        "figure projected to reach 68% by 2050 according to United Nations estimates. "
        "This demographic shift has been accompanied by a growing body of evidence linking "
        "urban living to elevated rates of mental health disorders, including depression, "
        "anxiety, and chronic stress."
    )
    run.font.size = Pt(11)

    intro_p2 = doc.add_paragraph()
    intro_p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = intro_p2.add_run(
        "Green spaces\u2014defined here as publicly accessible parks, community gardens, "
        "tree-lined corridors, and natural reserves within urban boundaries\u2014have "
        "emerged as a potential moderating factor. Seminal work by Kaplan and Kaplan "
        "(1989) on Attention Restoration Theory proposed that natural environments "
        "facilitate cognitive recovery from mental fatigue, while Ulrich's Stress "
        "Reduction Theory (1991) posited that visual exposure to vegetation triggers "
        "measurable decreases in cortisol and sympathetic nervous system activation."
    )
    run.font.size = Pt(11)

    # --- Body text: Methodology ---
    meth_heading = doc.add_heading("2. Methodology", level=1)
    meth_heading.paragraph_format.space_before = Pt(18)

    meth_p1 = doc.add_paragraph()
    meth_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = meth_p1.add_run(
        "Participants were recruited through stratified random sampling from municipal "
        "registries in twelve cities: Boston, Chicago, Denver, Los Angeles, Miami, "
        "Minneapolis, New York, Philadelphia, Phoenix, Portland, San Francisco, and "
        "Seattle. Eligibility criteria required participants to be between 18 and 65 "
        "years of age, reside in their current dwelling for at least twelve months, "
        "and have no prior clinical diagnosis of psychotic spectrum disorders."
    )
    run.font.size = Pt(11)

    meth_p2 = doc.add_paragraph()
    meth_p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = meth_p2.add_run(
        "Green space proximity was measured using Geographic Information System (GIS) "
        "analysis, calculating the Euclidean distance from each participant's residential "
        "address to the nearest qualifying green space boundary. Participants were "
        "categorized into three distance bands: proximal (< 500 m), moderate (500 m \u2013 "
        "1.5 km), and distal (> 1.5 km). Mental health outcomes were assessed using "
        "the Generalized Anxiety Disorder 7-item scale (GAD-7), the Patient Health "
        "Questionnaire 9-item scale (PHQ-9), and the Warwick-Edinburgh Mental Well-being "
        "Scale (WEMWBS)."
    )
    run.font.size = Pt(11)

    # --- Body text: Results ---
    results_heading = doc.add_heading("3. Results", level=1)
    results_heading.paragraph_format.space_before = Pt(18)

    results_p1 = doc.add_paragraph()
    results_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = results_p1.add_run(
        "After controlling for socioeconomic status, age, gender, baseline mental health "
        "scores, and neighborhood walkability indices, multilevel regression analysis "
        "revealed that proximity to green spaces was significantly associated with lower "
        "GAD-7 scores (\u03b2 = \u22120.34, SE = 0.08, p < 0.001) and higher WEMWBS scores "
        "(\u03b2 = 0.41, SE = 0.09, p < 0.001). The PHQ-9 association was also significant "
        "but with a smaller effect size (\u03b2 = \u22120.22, SE = 0.07, p = 0.002). "
        "Subgroup analysis indicated that the green space effect was strongest among "
        "participants aged 25\u201345 residing in neighborhoods with population densities "
        "exceeding 10,000 persons per square kilometer."
    )
    run.font.size = Pt(11)

    results_p2 = doc.add_paragraph()
    results_p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = results_p2.add_run(
        "Longitudinal trajectory modeling showed that improvements in well-being scores "
        "among the proximal group stabilized after approximately 18 months of exposure, "
        "suggesting a saturation effect. Notably, participants who relocated from distal "
        "to proximal zones during the study period (n = 137) exhibited a mean GAD-7 "
        "reduction of 3.2 points within 12 months of relocation, compared to a 0.4-point "
        "reduction in the non-relocating distal group over the same interval."
    )
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
