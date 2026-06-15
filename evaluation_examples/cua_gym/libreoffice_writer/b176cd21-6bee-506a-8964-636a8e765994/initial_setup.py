"""
Initial Setup: Create a comparison table document with uniform 0.5pt single borders.
Task ID: writer_tm_047
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_047'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell using XML.
    Each border arg is a dict with keys: sz, val, color, space.
    val: 'single', 'double', 'none', etc.
    sz: size in eighths of a point (e.g., 4 = 0.5pt, 8 = 1pt)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    else:
        # Remove existing to rebuild
        tcPr.remove(tcBorders)
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)

    for edge, props in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if props is not None:
            border_el = parse_xml(
                f'<w:{edge} {nsdecls("w")} '
                f'w:val="{props["val"]}" '
                f'w:sz="{props["sz"]}" '
                f'w:space="0" '
                f'w:color="{props["color"]}"/>'
            )
            tcBorders.append(border_el)


def create_initial():
    doc = Document()

    # Add title
    heading = doc.add_heading('Smartphone Comparison 2025', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        'The following table compares key specifications of the top-selling '
        'smartphones released in the first quarter of 2025. Data sourced from '
        'manufacturer press releases and independent benchmark testing.'
    )

    # Create a 4-row x 6-column comparison table
    # Row 0: headers, Rows 1-3: data
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'  # Gives uniform single-line borders

    # Headers
    headers = ['Feature', 'Galaxy S25 Ultra', 'iPhone 16 Pro', 'Pixel 9 Pro', 'OnePlus 13', 'Xiaomi 15 Pro']
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Data rows
    data = [
        ['Display Size', '6.9"', '6.3"', '6.3"', '6.82"', '6.73"'],
        ['Battery (mAh)', '5,000', '3,582', '4,700', '6,000', '5,400'],
        ['Base Price (USD)', '$1,299', '$1,099', '$999', '$899', '$749'],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if col_idx == 0:
                run.bold = True

    # Apply uniform 0.5pt single borders to all cells explicitly
    border_single_half = {'val': 'single', 'sz': '4', 'color': '000000'}
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(
                cell,
                top=border_single_half,
                bottom=border_single_half,
                left=border_single_half,
                right=border_single_half,
            )

    # Add a note below the table
    doc.add_paragraph('')
    note = doc.add_paragraph('Note: Prices reflect suggested retail pricing at launch. '
                             'Actual prices may vary by region and carrier.')
    note_run = note.runs[0]
    note_run.font.size = Pt(9)
    note_run.font.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
