"""
Reward Script: Apply double-line outer borders (1pt) and single-line inner borders (0.5pt) to comparison table
Task ID: writer_tm_047
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Outer borders are double-line style
  Component 2 (0.35): Outer borders have correct size (sz=8, i.e. 1pt)
  Component 3 (0.15): Inner borders remain single-line style
  Component 4 (0.15): Cell content unchanged
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_047'

# Expected table content for data integrity check
EXPECTED_CONTENT = [
    ["Feature", "Galaxy S25 Ultra", "iPhone 16 Pro", "Pixel 9 Pro", "OnePlus 13", "Xiaomi 15 Pro"],
    ["Display Size", '6.9"', '6.3"', '6.3"', '6.82"', '6.73"'],
    ["Battery (mAh)", "5,000", "3,582", "4,700", "6,000", "5,400"],
    ["Base Price (USD)", "$1,299", "$1,099", "$999", "$899", "$749"],
]


def get_cell_border(cell, side):
    """Get border properties for a cell side (top/bottom/left/right).
    Returns (val, sz) tuple or (None, None) if not found."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    tc = cell._tc
    tcPr = tc.find('.//w:tcPr', ns)
    if tcPr is None:
        return None, None
    tcBorders = tcPr.find('w:tcBorders', ns)
    if tcBorders is None:
        return None, None
    border_el = tcBorders.find(f'w:{side}', ns)
    if border_el is None:
        return None, None
    val = border_el.attrib.get(f'{{{ns["w"]}}}val', border_el.attrib.get('val'))
    sz = border_el.attrib.get(f'{{{ns["w"]}}}sz', border_el.attrib.get('sz'))
    # Handle namespace-prefixed attribs
    for k, v in border_el.attrib.items():
        if k.endswith('}val') or k == 'val':
            val = v
        if k.endswith('}sz') or k == 'sz':
            sz = v
    return val, sz


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: table exists with correct dimensions
    if len(doc.tables) < 1:
        print("FAIL: No table found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    if num_rows != 4 or num_cols != 6:
        print(f"FAIL: Expected 4x6 table, found {num_rows}x{num_cols}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Outer borders are double-line style (0.35 points)
    # Outer borders = top edge of row 0, bottom edge of last row,
    #                 left edge of col 0, right edge of last col
    try:
        outer_double_count = 0
        outer_total = 0

        for ri in range(num_rows):
            for ci in range(num_cols):
                cell = table.cell(ri, ci)
                # Top edge of table (row 0)
                if ri == 0:
                    val, sz = get_cell_border(cell, 'top')
                    outer_total += 1
                    if val == 'double':
                        outer_double_count += 1
                # Bottom edge of table (last row)
                if ri == num_rows - 1:
                    val, sz = get_cell_border(cell, 'bottom')
                    outer_total += 1
                    if val == 'double':
                        outer_double_count += 1
                # Left edge of table (col 0)
                if ci == 0:
                    val, sz = get_cell_border(cell, 'left')
                    outer_total += 1
                    if val == 'double':
                        outer_double_count += 1
                # Right edge of table (last col)
                if ci == num_cols - 1:
                    val, sz = get_cell_border(cell, 'right')
                    outer_total += 1
                    if val == 'double':
                        outer_double_count += 1

        outer_ratio = outer_double_count / outer_total if outer_total > 0 else 0
        if outer_ratio >= 0.9:
            print(f"PASS: Component 1 -- Outer borders are double-line ({outer_double_count}/{outer_total}) (0.35 pts)")
            total_score += 0.35
        elif outer_ratio >= 0.5:
            partial = round(0.35 * outer_ratio, 2)
            print(f"PARTIAL: Component 1 -- Outer borders double-line {outer_double_count}/{outer_total} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 -- Outer borders double-line {outer_double_count}/{outer_total}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Outer borders have correct size sz=8 (1pt) (0.35 points)
    try:
        outer_sz_correct = 0
        outer_sz_total = 0

        for ri in range(num_rows):
            for ci in range(num_cols):
                cell = table.cell(ri, ci)
                if ri == 0:
                    val, sz = get_cell_border(cell, 'top')
                    outer_sz_total += 1
                    if sz is not None and int(sz) >= 6:  # sz=8 is 1pt; allow some tolerance
                        outer_sz_correct += 1
                if ri == num_rows - 1:
                    val, sz = get_cell_border(cell, 'bottom')
                    outer_sz_total += 1
                    if sz is not None and int(sz) >= 6:
                        outer_sz_correct += 1
                if ci == 0:
                    val, sz = get_cell_border(cell, 'left')
                    outer_sz_total += 1
                    if sz is not None and int(sz) >= 6:
                        outer_sz_correct += 1
                if ci == num_cols - 1:
                    val, sz = get_cell_border(cell, 'right')
                    outer_sz_total += 1
                    if sz is not None and int(sz) >= 6:
                        outer_sz_correct += 1

        sz_ratio = outer_sz_correct / outer_sz_total if outer_sz_total > 0 else 0
        if sz_ratio >= 0.9:
            print(f"PASS: Component 2 -- Outer border sizes correct ({outer_sz_correct}/{outer_sz_total}) (0.35 pts)")
            total_score += 0.35
        elif sz_ratio >= 0.5:
            partial = round(0.35 * sz_ratio, 2)
            print(f"PARTIAL: Component 2 -- Outer border sizes {outer_sz_correct}/{outer_sz_total} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- Outer border sizes correct {outer_sz_correct}/{outer_sz_total}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Inner borders remain single-line 0.5pt (0.15 points)
    # Inner borders = borders between cells that are NOT on the outer edge
    try:
        inner_single_count = 0
        inner_total = 0

        for ri in range(num_rows):
            for ci in range(num_cols):
                cell = table.cell(ri, ci)
                # Top border of non-first rows (inner horizontal)
                if ri > 0:
                    val, sz = get_cell_border(cell, 'top')
                    inner_total += 1
                    if val == 'single' and sz is not None and int(sz) <= 6:
                        inner_single_count += 1
                # Bottom border of non-last rows (inner horizontal)
                if ri < num_rows - 1:
                    val, sz = get_cell_border(cell, 'bottom')
                    inner_total += 1
                    if val == 'single' and sz is not None and int(sz) <= 6:
                        inner_single_count += 1
                # Left border of non-first cols (inner vertical)
                if ci > 0:
                    val, sz = get_cell_border(cell, 'left')
                    inner_total += 1
                    if val == 'single' and sz is not None and int(sz) <= 6:
                        inner_single_count += 1
                # Right border of non-last cols (inner vertical)
                if ci < num_cols - 1:
                    val, sz = get_cell_border(cell, 'right')
                    inner_total += 1
                    if val == 'single' and sz is not None and int(sz) <= 6:
                        inner_single_count += 1

        inner_ratio = inner_single_count / inner_total if inner_total > 0 else 0
        if inner_ratio >= 0.9:
            print(f"PASS: Component 3 -- Inner borders are single-line ({inner_single_count}/{inner_total}) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 -- Inner borders single-line {inner_single_count}/{inner_total}")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Cell content unchanged (0.15 points)
    # This checks that ONLY borders changed, not content
    # On initial_env, all borders are single/4, so outer borders won't be double -> components 1&2 fail
    # But content IS the same on initial. So we make this a compound check:
    #   content is intact AND at least one outer border is double
    try:
        content_match = True
        for ri in range(num_rows):
            for ci in range(num_cols):
                actual = table.cell(ri, ci).text.strip()
                expected = EXPECTED_CONTENT[ri][ci]
                if actual != expected:
                    content_match = False
                    print(f"  Cell({ri},{ci}): expected '{expected}', found '{actual}'")

        # Compound: content intact AND outer borders changed (at least 1 double border)
        has_any_double = False
        for ci in range(num_cols):
            val, _ = get_cell_border(table.cell(0, ci), 'top')
            if val == 'double':
                has_any_double = True
                break

        if content_match and has_any_double:
            print(f"PASS: Component 4 -- Content intact with border changes applied (0.15 pts)")
            total_score += 0.15
        elif content_match and not has_any_double:
            print(f"FAIL: Component 4 -- Content intact but no double borders found (no task change)")
        else:
            print(f"FAIL: Component 4 -- Cell content was modified")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
