"""
Initial Setup: Achievements document with standard bullet list (no image bullets)
Task ID: writer_list_033
Domain: libreoffice_writer

Creates:
  - /home/user/Desktop/achievements.docx — 6 items with standard round bullets
  - /home/user/Desktop/star_bullet.gif   — 10x10 GIF star image for agent to use
"""

import os
import shlex
import subprocess
import time
import struct
import zlib

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_list_033'
OUTPUT_DOC = f'{WORKDIR}/achievements.docx'
OUTPUT_GIF = f'{WORKDIR}/star_bullet.gif'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_star_gif():
    """Create a 10x10 GIF with a simple yellow star on transparent background."""
    # GIF89a format: 10x10 pixels, star shape
    # Build a simple 10x10 GIF with a star pattern
    # Colors: 0=transparent, 1=yellow (#FFD700)

    width, height = 10, 10

    # Star pattern for 10x10 grid (1=yellow star pixel, 0=transparent)
    star_pixels = [
        [0, 0, 0, 0, 1, 1, 0, 0, 0, 0],  # row 0
        [0, 0, 0, 0, 1, 1, 0, 0, 0, 0],  # row 1
        [0, 1, 0, 1, 1, 1, 1, 0, 1, 0],  # row 2
        [0, 0, 1, 1, 1, 1, 1, 1, 0, 0],  # row 3
        [1, 1, 1, 1, 1, 1, 1, 1, 1, 1],  # row 4
        [0, 1, 1, 1, 1, 1, 1, 1, 1, 0],  # row 5
        [0, 0, 1, 0, 1, 1, 0, 1, 0, 0],  # row 6
        [0, 1, 0, 0, 1, 1, 0, 0, 1, 0],  # row 7
        [1, 0, 0, 0, 0, 0, 0, 0, 0, 1],  # row 8
        [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # row 9
    ]

    # Build GIF manually
    # Header
    gif_data = bytearray()
    gif_data += b'GIF89a'

    # Logical Screen Descriptor: width, height, packed, bgcolor, aspect
    # packed: 0x91 = Global Color Table present, color resolution 2, not sorted, 2^(1+1)=4 colors
    gif_data += struct.pack('<H', width)   # width
    gif_data += struct.pack('<H', height)  # height
    gif_data += bytes([0x91, 0, 0])       # packed=GCT present, 4 colors; bgcolor idx=0; aspect=0

    # Global Color Table: 4 colors (need 2^(n+1) where n=1, so 4 colors)
    # Color 0: transparent (black placeholder)
    # Color 1: yellow #FFD700
    # Color 2: unused
    # Color 3: unused
    gif_data += bytes([0x00, 0x00, 0x00])  # color 0: black/transparent
    gif_data += bytes([0xFF, 0xD7, 0x00])  # color 1: yellow
    gif_data += bytes([0x00, 0x00, 0x00])  # color 2: unused
    gif_data += bytes([0x00, 0x00, 0x00])  # color 3: unused

    # Graphic Control Extension (for transparency)
    gif_data += bytes([0x21, 0xF9, 0x04])  # GCE block
    gif_data += bytes([0x01])              # packed: transparent color flag set
    gif_data += struct.pack('<H', 0)       # delay time
    gif_data += bytes([0x00])             # transparent color index = 0
    gif_data += bytes([0x00])             # block terminator

    # Image Descriptor
    gif_data += bytes([0x2C])             # image separator
    gif_data += struct.pack('<H', 0)      # left
    gif_data += struct.pack('<H', 0)      # top
    gif_data += struct.pack('<H', width)  # width
    gif_data += struct.pack('<H', height) # height
    gif_data += bytes([0x00])             # packed: no local color table, not interlaced

    # Image Data - LZW compressed
    # Flatten pixel data
    pixels = []
    for row in star_pixels:
        pixels.extend(row)

    # LZW minimum code size for 4-color palette = 2
    lzw_min = 2
    gif_data += bytes([lzw_min])

    # Simple LZW compression
    def lzw_compress(data, min_code_size):
        clear_code = 1 << min_code_size
        eoi_code = clear_code + 1
        code_size = min_code_size + 1
        max_code = (1 << code_size)

        code_table = {(i,): i for i in range(clear_code)}
        next_code = eoi_code + 1

        output_bits = []

        def emit(code, size):
            for i in range(size):
                output_bits.append((code >> i) & 1)

        emit(clear_code, code_size)

        index_buffer = (data[0],)
        for pixel in data[1:]:
            index_buffer_plus = index_buffer + (pixel,)
            if index_buffer_plus in code_table:
                index_buffer = index_buffer_plus
            else:
                emit(code_table[index_buffer], code_size)
                if next_code < 4096:
                    code_table[index_buffer_plus] = next_code
                    next_code += 1
                    if next_code > max_code and code_size < 12:
                        code_size += 1
                        max_code = (1 << code_size)
                elif next_code == 4096:
                    emit(clear_code, code_size)
                    code_size = min_code_size + 1
                    max_code = (1 << code_size)
                    code_table = {(i,): i for i in range(clear_code)}
                    next_code = eoi_code + 1
                index_buffer = (pixel,)

        emit(code_table[index_buffer], code_size)
        emit(eoi_code, code_size)

        # Pack bits into bytes (LSB first)
        result = bytearray()
        for i in range(0, len(output_bits), 8):
            byte = 0
            for j in range(8):
                if i + j < len(output_bits):
                    byte |= output_bits[i + j] << j
            result.append(byte)
        return bytes(result)

    compressed = lzw_compress(pixels, lzw_min)

    # Write in sub-blocks of max 255 bytes
    i = 0
    while i < len(compressed):
        chunk = compressed[i:i+255]
        gif_data += bytes([len(chunk)])
        gif_data += chunk
        i += 255
    gif_data += bytes([0x00])  # block terminator

    # GIF trailer
    gif_data += bytes([0x3B])

    with open(OUTPUT_GIF, 'wb') as f:
        f.write(gif_data)
    print(f'Star GIF created: {OUTPUT_GIF}')


def create_initial():
    """Create achievements.docx with 6 standard bullet items (no image bullets)."""
    from docx import Document
    from docx.shared import Pt

    # Ensure Desktop dir exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Add a title heading
    title = doc.add_heading('Company Achievements', level=1)

    # Add a brief intro paragraph
    intro = doc.add_paragraph('The following achievements represent our key accomplishments over the past year:')

    # Six achievement items as standard bullet list
    achievements = [
        'Completed ISO 9001 certification',
        'Won Best Workplace Award',
        'Achieved zero workplace incidents for 365 days',
        'Launched three new product lines',
        'Expanded to five international markets',
        'Published ten industry white papers',
    ]

    for item in achievements:
        doc.add_paragraph(item, style='List Bullet')

    doc.save(OUTPUT_DOC)
    print(f'Initial file created: {OUTPUT_DOC}')

    # Create the star GIF
    create_star_gif()

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT_DOC}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
