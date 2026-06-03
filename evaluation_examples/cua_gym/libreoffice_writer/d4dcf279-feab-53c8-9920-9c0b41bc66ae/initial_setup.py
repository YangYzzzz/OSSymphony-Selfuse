"""
Initial Setup: Create a raw screenplay scene document with default formatting.
Task ID: wrpara_032
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'wrpara_032'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(12)

    # Title line
    doc.add_paragraph('UNTITLED SCREENPLAY')
    doc.add_paragraph('')

    # Scene heading
    doc.add_paragraph('INT. CORPORATE CONFERENCE ROOM - MORNING')
    doc.add_paragraph('')

    # Stage direction 1
    doc.add_paragraph('(The conference room is dimly lit. A long oak table dominates the center, surrounded by leather chairs. Rain streaks down the floor-to-ceiling windows.)')
    doc.add_paragraph('')

    # Character name + dialogue 1
    doc.add_paragraph('ALICE')
    doc.add_paragraph('We need to talk about the merger. The board is expecting a final recommendation by Friday, and we still have three unresolved issues on the due diligence checklist.')
    doc.add_paragraph('')

    # Stage direction 2
    doc.add_paragraph('(Bob slides a manila folder across the table, his expression tense.)')
    doc.add_paragraph('')

    # Character name + dialogue 2
    doc.add_paragraph('BOB')
    doc.add_paragraph("I've been reviewing the financials all weekend. The target company understated their liabilities by nearly twelve percent. That alone should give us pause.")
    doc.add_paragraph('')

    # Character name + dialogue 3
    doc.add_paragraph('ALICE')
    doc.add_paragraph('Twelve percent is within the range we anticipated. The real concern is their intellectual property portfolio. Half of those patents expire within eighteen months.')
    doc.add_paragraph('')

    # Stage direction 3
    doc.add_paragraph('(Carol enters through the glass door, carrying a stack of printed reports. She sets them down at the head of the table.)')
    doc.add_paragraph('')

    # Character name + dialogue 4
    doc.add_paragraph('CAROL')
    doc.add_paragraph('Sorry I am late. Legal just finished the antitrust analysis. We have a green light from regulatory counsel, but there is a caveat regarding the European subsidiary.')
    doc.add_paragraph('')

    # Character name + dialogue 5
    doc.add_paragraph('BOB')
    doc.add_paragraph('What kind of caveat? The last thing we need is a cross-border complication delaying the timeline.')
    doc.add_paragraph('')

    # Stage direction 4
    doc.add_paragraph('(Carol opens one of the reports and points to a highlighted section.)')
    doc.add_paragraph('')

    # Character name + dialogue 6
    doc.add_paragraph('CAROL')
    doc.add_paragraph('The European subsidiary has an ongoing dispute with a local supplier. It is not a dealbreaker, but it could add six to eight weeks to the integration schedule if we do not address it upfront.')
    doc.add_paragraph('')

    # Character name + dialogue 7
    doc.add_paragraph('ALICE')
    doc.add_paragraph('Then we address it upfront. Bob, can your team draft a resolution framework by Wednesday? Carol, loop in outside counsel for the European angle. We are not letting this slip.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
