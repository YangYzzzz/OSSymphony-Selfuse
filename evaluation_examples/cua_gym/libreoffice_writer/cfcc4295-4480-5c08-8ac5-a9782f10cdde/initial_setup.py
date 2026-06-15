"""
Initial Setup: Sort inventory table by Category ascending, then Price descending
Task ID: writer_tm_048
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_048'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title
    heading = doc.add_heading('Warehouse Inventory Report', level=1)

    # Add a brief intro paragraph
    doc.add_paragraph(
        'Below is the current inventory listing for Q1 2025. '
        'Please review and update as needed.'
    )

    # Create 5x21 table (1 header + 20 data rows) - UNSORTED
    table = doc.add_table(rows=21, cols=5)
    table.style = 'Table Grid'

    # Headers
    headers = ['ID', 'Name', 'Category', 'Price', 'Stock']
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    # Data rows - deliberately UNSORTED to make the task meaningful
    data = [
        ['INV-001', 'Standing Desk Pro',        'Furniture',        749.99,  12],
        ['INV-002', 'Wireless Keyboard K380',   'Electronics',      49.99,   85],
        ['INV-003', 'Manila Folders (100pk)',    'Office Supplies',  12.49,  200],
        ['INV-004', 'USB-C Docking Station',    'Electronics',      189.99,  34],
        ['INV-005', 'Ergonomic Mesh Chair',     'Furniture',        529.00,  18],
        ['INV-006', 'Ballpoint Pens (50pk)',    'Office Supplies',   8.99,  350],
        ['INV-007', '27" 4K Monitor',           'Electronics',      449.99,  22],
        ['INV-008', 'Bookshelf Oak 5-Tier',     'Furniture',        189.50,  15],
        ['INV-009', 'Sticky Notes Assorted',    'Office Supplies',   5.49,  500],
        ['INV-010', 'Noise Cancelling Headset', 'Electronics',      299.99,  40],
        ['INV-011', 'Filing Cabinet 3-Drawer',  'Furniture',        274.00,   8],
        ['INV-012', 'Laser Printer Toner',      'Office Supplies',  64.99,   60],
        ['INV-013', 'Webcam HD 1080p',          'Electronics',       79.99,  55],
        ['INV-014', 'Conference Table Round',   'Furniture',       1249.00,   3],
        ['INV-015', 'Paper Clips Jumbo (500)',  'Office Supplies',   3.29,  800],
        ['INV-016', 'Portable SSD 1TB',         'Electronics',      109.99,  70],
        ['INV-017', 'Whiteboard 48x36',         'Office Supplies',  89.99,   25],
        ['INV-018', 'Desk Lamp LED Adjustable', 'Furniture',         64.50,  45],
        ['INV-019', 'Mechanical Keyboard RGB',  'Electronics',      139.99,  30],
        ['INV-020', 'Binder Clips Large (24pk)','Office Supplies',   6.79,  420],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            if isinstance(val, float):
                cell.text = f'{val:.2f}'
            elif isinstance(val, int) and col_idx == 4:
                cell.text = str(val)
            else:
                cell.text = str(val)

    # Add a closing paragraph
    doc.add_paragraph('')
    doc.add_paragraph(
        'Last updated: March 28, 2025 | Prepared by: Operations Team'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
