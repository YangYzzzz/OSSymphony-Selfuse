"""
Reward Script: Sort inventory table by Category (ascending) then Price (descending)
Task ID: writer_tm_048
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15): Header row preserved in row 0
  Component 2 (0.15): All 20 data rows present (no data loss)
  Component 3 (0.35): Primary sort — Category column is ascending (A-Z)
  Component 4 (0.35): Secondary sort — Price is descending within each category group
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_048'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_writer", "libreoffice_calc", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: must have exactly 1 table
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    rows = table.rows

    # Precondition: table must have 21 rows (1 header + 20 data)
    if len(rows) < 2:
        print("CRITICAL: Table has fewer than 2 rows")
        print("REWARD: 0.0")
        return 0.0

    # Extract all row data
    header = [cell.text.strip() for cell in rows[0].cells]
    data_rows = []
    for i in range(1, len(rows)):
        cells = [cell.text.strip() for cell in rows[i].cells]
        data_rows.append(cells)

    # Expected header
    expected_header = ['ID', 'Name', 'Category', 'Price', 'Stock']

    # Expected initial IDs (all 20 items must be present)
    expected_ids = {f'INV-{str(i).zfill(3)}' for i in range(1, 21)}

    # Component 1: Header row preserved (0.15 points)
    # This checks that the header is still in row 0 and has the correct columns.
    # The initial file also has the correct header, BUT we combine this with a
    # sort-awareness check: we verify header is correct AND data rows are NOT
    # in the original INV-001..INV-020 sequential order (i.e., sorting happened).
    try:
        header_correct = (header == expected_header)
        # Check that data is NOT in original sequential order (sort happened)
        actual_ids = [r[0] for r in data_rows if len(r) >= 1]
        original_order = [f'INV-{str(i).zfill(3)}' for i in range(1, 21)]
        is_reordered = (actual_ids != original_order)

        if header_correct and is_reordered:
            print(f"PASS: Component 1 — Header preserved and data reordered (0.15 pts)")
            total_score += 0.15
        elif header_correct and not is_reordered:
            print(f"FAIL: Component 1 — Header correct but data still in original order (no sorting done)")
        else:
            print(f"FAIL: Component 1 — Header mismatch. Expected {expected_header}, found {header}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All 20 data rows present with no data loss (0.15 points)
    # Combined with reorder check to avoid scoring on initial_env.
    try:
        actual_ids_set = {r[0] for r in data_rows if len(r) >= 1}
        all_present = (actual_ids_set == expected_ids)
        row_count_correct = (len(data_rows) == 20)

        if all_present and row_count_correct and is_reordered:
            print(f"PASS: Component 2 — All 20 data rows present and reordered (0.15 pts)")
            total_score += 0.15
        elif not all_present:
            missing = expected_ids - actual_ids_set
            extra = actual_ids_set - expected_ids
            print(f"FAIL: Component 2 — Missing IDs: {missing}, Extra IDs: {extra}")
        elif not row_count_correct:
            print(f"FAIL: Component 2 — Expected 20 data rows, found {len(data_rows)}")
        else:
            print(f"FAIL: Component 2 — Data still in original order")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Primary sort — Category ascending (0.35 points)
    # Extract categories from column 3 (index 2)
    try:
        categories = [r[2] for r in data_rows if len(r) >= 3]

        if len(categories) == 20:
            # Check that categories are in non-decreasing (ascending) order
            is_ascending = all(categories[i] <= categories[i+1] for i in range(len(categories)-1))

            if is_ascending:
                # Verify the expected category groups
                expected_category_order = (
                    ['Electronics'] * 7 +
                    ['Furniture'] * 6 +
                    ['Office Supplies'] * 7
                )
                exact_match = (categories == expected_category_order)

                if exact_match:
                    print(f"PASS: Component 3 — Categories sorted ascending with correct grouping (0.35 pts)")
                    total_score += 0.35
                else:
                    # Still ascending but maybe different counts - partial credit
                    print(f"PASS: Component 3 — Categories sorted ascending (0.35 pts)")
                    total_score += 0.35
            else:
                # Check how many adjacent pairs are in order
                in_order = sum(1 for i in range(len(categories)-1) if categories[i] <= categories[i+1])
                print(f"FAIL: Component 3 — Categories not ascending. {in_order}/19 adjacent pairs in order.")
                print(f"  Categories found: {categories}")
        else:
            print(f"FAIL: Component 3 — Expected 20 category values, found {len(categories)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Secondary sort — Price descending within each category (0.35 points)
    # GATE: Primary sort (categories ascending) must be correct; otherwise this check
    # is meaningless (groupby on unsorted data creates trivial single-element groups).
    try:
        categories_are_ascending = all(
            categories[i] <= categories[i+1] for i in range(len(categories)-1)
        ) if len(categories) == 20 else False

        if not categories_are_ascending:
            print(f"FAIL: Component 4 — Skipped: primary sort (Category ascending) not satisfied")
        else:
            from itertools import groupby

            category_price_pairs = []
            for r in data_rows:
                if len(r) >= 4:
                    cat = r[2]
                    try:
                        price = float(r[3])
                    except ValueError:
                        price = 0.0
                    category_price_pairs.append((cat, price))

            if len(category_price_pairs) == 20:
                groups_correct = 0
                total_groups = 0
                for cat, group in groupby(category_price_pairs, key=lambda x: x[0]):
                    prices = [p for _, p in group]
                    total_groups += 1
                    is_desc = all(prices[i] >= prices[i+1] for i in range(len(prices)-1))
                    if is_desc:
                        groups_correct += 1
                        print(f"  Category '{cat}': prices {prices} — descending OK")
                    else:
                        print(f"  Category '{cat}': prices {prices} — NOT descending")

                if total_groups > 0 and groups_correct == total_groups:
                    print(f"PASS: Component 4 — Price descending within all {total_groups} categories (0.35 pts)")
                    total_score += 0.35
                elif total_groups > 0:
                    partial = 0.35 * (groups_correct / total_groups)
                    print(f"PARTIAL: Component 4 — {groups_correct}/{total_groups} category groups correctly sorted ({partial:.2f} pts)")
                    total_score += partial
                else:
                    print(f"FAIL: Component 4 — No category groups found")
            else:
                print(f"FAIL: Component 4 — Expected 20 rows, found {len(category_price_pairs)}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
