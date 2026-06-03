"""
Initial Setup: Contract document with 4 sections, all unprotected
Task ID: writer_biz_072
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_072'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # -- Document Title --
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    parties = doc.add_paragraph()
    parties.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = parties.add_run('Between Meridian Consulting Group, LLC and Apex Technologies Inc.')
    run.font.size = Pt(11)
    run.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Effective Date: March 15, 2025')
    run.font.size = Pt(11)

    doc.add_paragraph('')  # spacer

    # ===== SECTION 1: SCOPE OF SERVICES =====
    h1 = doc.add_heading('Section 1: Scope of Services', level=1)

    p = doc.add_paragraph()
    run = p.add_run('1.1 Service Description')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'Meridian Consulting Group, LLC ("Consultant") agrees to provide the following '
        'professional services to Apex Technologies Inc. ("Client"):'
    )

    doc.add_paragraph('Strategic technology infrastructure assessment and optimization', style='List Bullet')
    doc.add_paragraph('Cloud migration planning and implementation support', style='List Bullet')
    doc.add_paragraph('Cybersecurity audit and remediation roadmap development', style='List Bullet')
    doc.add_paragraph('Staff training and knowledge transfer programs', style='List Bullet')
    doc.add_paragraph('Quarterly performance review and reporting', style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('1.2 Deliverables')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'The Consultant shall deliver the following within the agreed timelines:'
    )

    # Deliverables table
    table1 = doc.add_table(rows=6, cols=3)
    table1.style = 'Table Grid'
    headers = ['Deliverable', 'Due Date', 'Status']
    for i, h in enumerate(headers):
        cell = table1.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    deliverables = [
        ['Infrastructure Assessment Report', 'April 30, 2025', 'Pending'],
        ['Cloud Migration Blueprint', 'June 15, 2025', 'Pending'],
        ['Security Audit Findings', 'August 1, 2025', 'Pending'],
        ['Training Curriculum Package', 'September 15, 2025', 'Pending'],
        ['Final Performance Summary', 'December 1, 2025', 'Pending'],
    ]
    for r, row_data in enumerate(deliverables, 1):
        for c, val in enumerate(row_data):
            table1.cell(r, c).text = val

    doc.add_paragraph('')  # spacer

    # ===== SECTION 2: COMPENSATION AND PAYMENT =====
    h2 = doc.add_heading('Section 2: Compensation and Payment', level=1)

    p = doc.add_paragraph()
    run = p.add_run('2.1 Fee Structure')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'The Client agrees to compensate the Consultant according to the following fee schedule:'
    )

    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    fee_headers = ['Service Category', 'Rate', 'Estimated Hours']
    for i, h in enumerate(fee_headers):
        cell = table2.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    fees = [
        ['Senior Consultant', '$275/hour', '320'],
        ['Technical Specialist', '$225/hour', '480'],
        ['Project Manager', '$200/hour', '160'],
        ['Administrative Support', '$125/hour', '80'],
    ]
    for r, row_data in enumerate(fees, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    doc.add_paragraph('')

    p = doc.add_paragraph()
    run = p.add_run('2.2 Payment Terms')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'Invoices shall be submitted monthly on the first business day following the service period. '
        'Payment is due within thirty (30) calendar days of invoice receipt. Late payments shall '
        'accrue interest at a rate of 1.5% per month on the outstanding balance.'
    )

    p = doc.add_paragraph()
    run = p.add_run('2.3 Expense Reimbursement')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'Reasonable travel, accommodation, and incidental expenses incurred in the performance '
        'of services shall be reimbursed upon submission of itemized receipts. Pre-approval is '
        'required for any single expense exceeding $500.00.'
    )

    # ===== SECTION 3: CONFIDENTIALITY AND INTELLECTUAL PROPERTY =====
    h3 = doc.add_heading('Section 3: Confidentiality and Intellectual Property', level=1)

    p = doc.add_paragraph()
    run = p.add_run('3.1 Confidential Information')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'Both parties acknowledge that during the course of this engagement, they may receive '
        'or have access to confidential and proprietary information belonging to the other party. '
        'Such information includes, but is not limited to, trade secrets, client lists, financial '
        'data, business strategies, technical specifications, and software source code.'
    )

    p = doc.add_paragraph()
    run = p.add_run('3.2 Non-Disclosure Obligations')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'Each party agrees to: (a) hold all Confidential Information in strict confidence; '
        '(b) not disclose Confidential Information to any third party without prior written '
        'consent; (c) use Confidential Information solely for purposes related to this Agreement; '
        'and (d) return or destroy all Confidential Information upon termination of this Agreement.'
    )

    p = doc.add_paragraph()
    run = p.add_run('3.3 Intellectual Property Rights')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'All work product, deliverables, and materials created by the Consultant specifically for '
        'the Client under this Agreement shall be considered "work made for hire" and shall become '
        'the exclusive property of the Client upon full payment. The Consultant retains ownership '
        'of pre-existing tools, methodologies, and frameworks used in the delivery of services.'
    )

    p = doc.add_paragraph()
    run = p.add_run('3.4 Survival')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'The obligations set forth in this Section 3 shall survive the termination or expiration '
        'of this Agreement for a period of five (5) years.'
    )

    # ===== SECTION 4: AMENDMENTS =====
    h4 = doc.add_heading('Section 4: Amendments', level=1)

    p = doc.add_paragraph()
    run = p.add_run('4.1 Modification Procedure')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'This Agreement may be amended or modified only by a written instrument signed by both '
        'parties. Any proposed amendments shall be submitted in writing to the other party at '
        'least fifteen (15) business days prior to the desired effective date.'
    )

    p = doc.add_paragraph()
    run = p.add_run('4.2 Pending Amendments')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'The following amendments are currently under review by the legal team:'
    )

    doc.add_paragraph('Amendment A: Extension of service period through Q2 2026', style='List Bullet')
    doc.add_paragraph('Amendment B: Addition of data analytics consulting services', style='List Bullet')
    doc.add_paragraph('Amendment C: Revised rate schedule effective January 2026', style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('4.3 Amendment Log')
    run.bold = True
    run.font.size = Pt(11)

    table3 = doc.add_table(rows=4, cols=4)
    table3.style = 'Table Grid'
    log_headers = ['Amendment ID', 'Description', 'Proposed By', 'Status']
    for i, h in enumerate(log_headers):
        cell = table3.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    amendments = [
        ['AMD-001', 'Service period extension', 'Client Legal', 'Under Review'],
        ['AMD-002', 'Analytics services addition', 'Consultant', 'Draft'],
        ['AMD-003', 'Rate schedule revision', 'Joint Committee', 'Pending Approval'],
    ]
    for r, row_data in enumerate(amendments, 1):
        for c, val in enumerate(row_data):
            table3.cell(r, c).text = val

    doc.add_paragraph('')

    # Signature block
    sig = doc.add_paragraph()
    sig.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = sig.add_run('_' * 40 + '\n')
    run.font.size = Pt(11)
    run = sig.add_run('Authorized Signature — Meridian Consulting Group, LLC\n\n')
    run.font.size = Pt(10)
    run = sig.add_run('_' * 40 + '\n')
    run.font.size = Pt(11)
    run = sig.add_run('Authorized Signature — Apex Technologies Inc.')
    run.font.size = Pt(10)

    # NO section protection applied — entire document is unprotected
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
