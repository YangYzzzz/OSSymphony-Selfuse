"""
Reward Script: Insert today's date field at the top right of termination letter (MM/DD/YYYY)
Task ID: writer_hr_022
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): A DATE field code exists in a paragraph near the top of the document
  Component 2 (0.3): The DATE field uses MM/DD/YYYY format
  Component 3 (0.3): The paragraph containing the date field is right-aligned
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_022'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            import time
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # We look for a date field in the first few paragraphs (top of document).
    # The task says "at the top" so we check the first 5 paragraphs.
    date_field_para_idx = None
    date_field_instr = None

    for i, para in enumerate(doc.paragraphs[:5]):
        xml_str = para._element.xml
        if 'fldChar' in xml_str and 'instrText' in xml_str:
            # Extract instrText content
            instr_matches = re.findall(r'instrText[^>]*>([^<]+)<', xml_str)
            for instr in instr_matches:
                if 'DATE' in instr.upper():
                    date_field_para_idx = i
                    date_field_instr = instr
                    break
        if date_field_para_idx is not None:
            break

    # Also accept if the agent typed a date string (not a field) in MM/DD/YYYY format
    # at the top with right alignment — but field is preferred.
    # We check for a date-like text in first 5 paragraphs as a fallback.
    date_text_para_idx = None
    if date_field_para_idx is None:
        for i, para in enumerate(doc.paragraphs[:5]):
            text = para.text.strip()
            # Match MM/DD/YYYY pattern
            if re.match(r'^\d{2}/\d{2}/\d{4}$', text):
                date_text_para_idx = i
                break

    # Component 1: A DATE field exists near the top of the document (0.4 points)
    try:
        if date_field_para_idx is not None:
            print(f"PASS: Component 1 — DATE field found in paragraph {date_field_para_idx} (0.4 pts)")
            total_score += 0.4
        elif date_text_para_idx is not None:
            # Partial credit: date text exists but not as a field code
            print(f"PARTIAL: Component 1 — Date text found in paragraph {date_text_para_idx} but not as a field code (0.2 pts)")
            total_score += 0.2
        else:
            print("FAIL: Component 1 — No DATE field or date text found in top 5 paragraphs")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: The DATE field uses MM/DD/YYYY format (0.3 points)
    try:
        if date_field_instr is not None:
            # Check for MM/DD/YYYY format in the instrText
            # Typical: DATE \@ "MM/DD/YYYY"
            if 'MM/DD/YYYY' in date_field_instr or 'MM\\/DD\\/YYYY' in date_field_instr:
                print(f"PASS: Component 2 — DATE field format is MM/DD/YYYY (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — DATE field format is not MM/DD/YYYY, found: {repr(date_field_instr)}")
        elif date_text_para_idx is not None:
            # Check that the typed text matches MM/DD/YYYY
            text = doc.paragraphs[date_text_para_idx].text.strip()
            if re.match(r'^\d{2}/\d{2}/\d{4}$', text):
                print(f"PARTIAL: Component 2 — Date text {text} matches MM/DD/YYYY format but is not a field (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 2 — Date text does not match MM/DD/YYYY format: {text}")
        else:
            print("FAIL: Component 2 — No date field found to check format")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: The date paragraph is right-aligned (0.3 points)
    try:
        target_idx = date_field_para_idx if date_field_para_idx is not None else date_text_para_idx
        if target_idx is not None:
            para = doc.paragraphs[target_idx]
            alignment = para.paragraph_format.alignment
            if alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                print(f"PASS: Component 3 — Date paragraph is right-aligned (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 3 — Date paragraph alignment is {alignment}, expected RIGHT (2)")
        else:
            print("FAIL: Component 3 — No date paragraph found to check alignment")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state('libreoffice_writer')
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
