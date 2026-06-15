"""
Initial Setup: Termination letter with no date field
Task ID: writer_hr_022
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_022'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Company letterhead
    company = doc.add_paragraph()
    company.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    company.paragraph_format.space_after = Pt(0)
    run = company.add_run("Meridian Technologies, Inc.")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"

    addr1 = doc.add_paragraph()
    addr1.paragraph_format.space_after = Pt(0)
    addr1.paragraph_format.space_before = Pt(0)
    r = addr1.add_run("4200 Innovation Drive, Suite 300")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    addr2 = doc.add_paragraph()
    addr2.paragraph_format.space_after = Pt(0)
    addr2.paragraph_format.space_before = Pt(0)
    r = addr2.add_run("San Jose, CA 95134")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    phone = doc.add_paragraph()
    phone.paragraph_format.space_after = Pt(12)
    phone.paragraph_format.space_before = Pt(0)
    r = phone.add_run("Phone: (408) 555-7200")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # Blank line (no date here -- task requires agent to insert date)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)

    # Recipient info
    recipient_lines = [
        "Jennifer A. Marshall",
        "Senior Software Engineer",
        "Engineering Department",
        "Employee ID: MT-2019-0847",
    ]
    for i, line in enumerate(recipient_lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.name = "Calibri"
        if i == 0:
            r.bold = True

    # Spacing before body
    spacer2 = doc.add_paragraph()
    spacer2.paragraph_format.space_after = Pt(6)

    # Subject line
    subject = doc.add_paragraph()
    subject.paragraph_format.space_after = Pt(12)
    r = subject.add_run("RE: Notice of Termination of Employment")
    r.bold = True
    r.underline = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # Salutation
    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(6)
    r = salutation.add_run("Dear Ms. Marshall,")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # Body paragraphs
    body_texts = [
        "This letter serves as formal notification that your employment with Meridian Technologies, Inc. "
        "is terminated effective March 28, 2026. This decision has been made following a thorough review "
        "process conducted by the Human Resources department in consultation with your direct supervisor, "
        "David Chen, Director of Engineering.",

        "As discussed during our meeting on March 14, 2026, the termination is based on continued "
        "performance deficiencies that were documented in your Performance Improvement Plan (PIP) "
        "initiated on January 6, 2026. Despite the support and resources provided during the 60-day "
        "improvement period, the required performance benchmarks were not satisfactorily met.",

        "Please review the following details regarding your separation:",
    ]
    for text in body_texts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # Bullet list of separation details
    bullet_items = [
        "Final paycheck, including accrued but unused vacation days (12 days), will be issued on the next regular pay date, April 3, 2026.",
        "Your company-sponsored health insurance coverage will remain active through April 30, 2026. COBRA continuation information will be mailed to your home address within 14 business days.",
        "You are entitled to a severance package equivalent to four (4) weeks of base salary ($7,692.30), contingent upon execution of the enclosed Separation Agreement and General Release.",
        "All company property, including your laptop (Asset Tag MT-L4521), access badge, and parking pass, must be returned to HR by close of business on your last working day.",
        "Your 401(k) account with Fidelity Investments will remain accessible. You may roll over your vested balance ($34,218.50) to a qualified plan of your choice.",
    ]
    for item in bullet_items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # Closing paragraphs
    closing_texts = [
        "Please note that your obligations under the Non-Disclosure Agreement (NDA) signed on August 12, "
        "2019, and the Non-Compete clause in your employment contract remain in effect for a period of "
        "twelve (12) months following your termination date.",

        "We encourage you to contact our Employee Assistance Program (EAP) at (800) 555-4357 if you "
        "need support during this transition. Additionally, our outplacement services partner, "
        "CareerBridge Solutions, will be available to assist you with your job search for up to 90 days.",

        "If you have any questions regarding this letter or the terms of your separation, please do not "
        "hesitate to contact me directly at (408) 555-7215 or via email at r.torres@meridiantech.com.",
    ]
    for text in closing_texts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # Sign-off
    spacer3 = doc.add_paragraph()
    spacer3.paragraph_format.space_after = Pt(6)

    sincerely = doc.add_paragraph()
    sincerely.paragraph_format.space_after = Pt(24)
    r = sincerely.add_run("Sincerely,")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    sig_lines = [
        ("Rachel Torres", True),
        ("Vice President, Human Resources", False),
        ("Meridian Technologies, Inc.", False),
    ]
    for text, bold in sig_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # Enclosure note
    spacer4 = doc.add_paragraph()
    spacer4.paragraph_format.space_after = Pt(6)

    enc = doc.add_paragraph()
    r = enc.add_run("Enclosures: Separation Agreement and General Release, COBRA Information Packet, "
                     "Outplacement Services Brochure")
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = "Calibri"

    # CC line
    cc = doc.add_paragraph()
    cc.paragraph_format.space_before = Pt(6)
    r = cc.add_run("cc: David Chen, Director of Engineering; Legal Department")
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
