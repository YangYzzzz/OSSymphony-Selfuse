"""
Reward Script: Newsletter Print Layout Setup
Task ID: writer_mktg_006
Domain: libreoffice_writer
Scoring:
  Component 1: Page margins set to 0.75 inches on all sides (0.25 pts)
  Component 2: Two-column layout in body section with ~0.3in gap (0.30 pts)
  Component 3: Column separator line enabled (0.20 pts)
  Component 4: Full-width header section (first 2 paragraphs + empty para) (0.15 pts)
  Component 5: Article headings are 14pt bold (0.10 pts)
Total: 1.0
"""

import os
from docx import Document
from docx.shared import Inches

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_mktg_006'
FILE_NAME = 'q1_client_newsletter.docx'
FILE_PATH = os.path.join(WORKDIR, FILE_NAME)

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Tolerance for margin/spacing checks (in inches)
MARGIN_TOL = 0.05


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: Page margins set to 0.75 inches on all sides (0.25 points)
    # Initial: 1.0 inch margins. Task requires 0.75 inch on all sides.
    # -------------------------------------------------------------------------
    try:
        all_margins_correct = True
        margin_details = []
        for i, section in enumerate(doc.sections):
            # Use raw EMU values for safety; 0.75 inches = 685800 EMU
            target_emu = int(Inches(0.75))
            tol_emu = int(Inches(MARGIN_TOL))
            margins = {
                'left': section.left_margin,
                'right': section.right_margin,
                'top': section.top_margin,
                'bottom': section.bottom_margin,
            }
            for side, val in margins.items():
                if val is None:
                    # Inherit — check if within tolerance
                    all_margins_correct = False
                    margin_details.append(f'Section {i} {side}: None (inherited)')
                elif abs(int(val) - target_emu) > tol_emu:
                    all_margins_correct = False
                    margin_details.append(
                        f'Section {i} {side}: {round(val.inches, 4)}in (expected ~0.75in)'
                    )
        if all_margins_correct:
            print("PASS: Component 1 — All section margins are 0.75in (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Margin issues: {margin_details}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Two-column layout in body section with ~0.3 inch gap (0.30 points)
    # Initial: single-column (cols element with no num attr, or num=1).
    # Golden: Section 1 has cols num=2 and space=432 twips (~0.3in).
    # We accept if ANY section has num=2 column layout.
    # -------------------------------------------------------------------------
    try:
        two_col_found = False
        gap_correct = False
        for section in doc.sections:
            sectPr = section._sectPr
            cols = sectPr.find(f'{{{NS}}}cols')
            if cols is not None:
                num_str = cols.get(f'{{{NS}}}num')
                if num_str is not None and int(num_str) == 2:
                    two_col_found = True
                    # Check gap: 0.3 inches = 432 twips (1 inch = 1440 twips)
                    space_str = cols.get(f'{{{NS}}}space')
                    if space_str is not None:
                        space_twips = int(space_str)
                        # Accept range: 360–576 twips (0.25–0.4 inches)
                        if 360 <= space_twips <= 576:
                            gap_correct = True
                            print(
                                f"PASS: Component 2 — Two-column layout found, gap={space_twips} twips"
                                f" (~{space_twips/1440:.3f}in) (0.30 pts)"
                            )
                            total_score += 0.30
                        else:
                            print(
                                f"FAIL: Component 2 — Two columns found but gap={space_twips} twips"
                                f" ({space_twips/1440:.3f}in), expected ~432 twips (0.3in)"
                            )
                    else:
                        # No explicit gap specified — default gap; partial pass for 2-col layout
                        print(
                            "FAIL: Component 2 — Two columns found but no explicit gap value set"
                        )
                    break
        if not two_col_found:
            print("FAIL: Component 2 — No two-column layout section found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Column separator line enabled (0.20 points)
    # Initial: no sep attribute. Golden: cols has sep=1.
    # -------------------------------------------------------------------------
    try:
        sep_found = False
        for section in doc.sections:
            sectPr = section._sectPr
            cols = sectPr.find(f'{{{NS}}}cols')
            if cols is not None:
                num_str = cols.get(f'{{{NS}}}num')
                sep_str = cols.get(f'{{{NS}}}sep')
                # sep="1" means separator line enabled
                if num_str is not None and int(num_str) == 2 and sep_str == '1':
                    sep_found = True
                    break
        if sep_found:
            print("PASS: Component 3 — Column separator line is enabled (0.20 pts)")
            total_score += 0.20
        else:
            print("FAIL: Component 3 — Column separator line not found in two-column section")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: Full-width header section (title + date in Section 0, body in Section 1) (0.15 pts)
    # Initial: 1 section. Golden: 2 sections where Section 0 is single-column (full-width header)
    # and Section 1 has the two-column body content.
    # We check: at least 2 sections exist AND first section has no 2-column layout (single-col header)
    # AND the section boundary falls after the first paragraph(s) (title+date area).
    # -------------------------------------------------------------------------
    try:
        has_multi_section = len(doc.sections) >= 2
        if has_multi_section:
            # Verify section 0 is single column (full-width)
            sec0_sectPr = doc.sections[0]._sectPr
            sec0_cols = sec0_sectPr.find(f'{{{NS}}}cols')
            sec0_is_single_col = True
            if sec0_cols is not None:
                sec0_num = sec0_cols.get(f'{{{NS}}}num')
                if sec0_num is not None and int(sec0_num) > 1:
                    sec0_is_single_col = False

            # Check that section break in paragraph exists near title area (para 0-5)
            section_break_para_idx = None
            for i, p in enumerate(doc.paragraphs):
                sectPr_in_para = p._element.find(f'.//{{{NS}}}pPr/{{{NS}}}sectPr')
                if sectPr_in_para is not None:
                    section_break_para_idx = i
                    break

            if sec0_is_single_col and section_break_para_idx is not None and section_break_para_idx <= 5:
                print(
                    f"PASS: Component 4 — Full-width header section detected. "
                    f"Section break after para {section_break_para_idx} (0.15 pts)"
                )
                total_score += 0.15
            elif not sec0_is_single_col:
                print("FAIL: Component 4 — Section 0 has multi-column layout (should be full-width)")
            elif section_break_para_idx is None:
                print("FAIL: Component 4 — No section break found inside paragraphs")
            else:
                print(
                    f"FAIL: Component 4 — Section break found at para {section_break_para_idx}, "
                    f"expected within first few paragraphs (title/date area)"
                )
        else:
            print("FAIL: Component 4 — Only one section found; need separate full-width header section")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -------------------------------------------------------------------------
    # Component 5: Article headings are 14pt bold (0.10 points)
    # Initial: headings are 12pt bold. Task asks to set them to 14pt bold.
    # We check the 6 article heading paragraphs (indices 3, 6, 9, 12, 15, 18).
    # -------------------------------------------------------------------------
    try:
        heading_indices = [3, 6, 9, 12, 15, 18]
        headings_ok = 0
        heading_total = 0
        for idx in heading_indices:
            if idx < len(doc.paragraphs):
                p = doc.paragraphs[idx]
                heading_total += 1
                for run in p.runs:
                    if run.text.strip():
                        size_pt = run.font.size.pt if run.font.size else None
                        is_bold = run.font.bold
                        # Accept 14pt (with small tolerance) and bold
                        if size_pt is not None and abs(size_pt - 14.0) < 0.5 and is_bold:
                            headings_ok += 1
                            break

        if heading_total > 0 and headings_ok == heading_total:
            print(
                f"PASS: Component 5 — All {headings_ok}/{heading_total} article headings are 14pt bold (0.10 pts)"
            )
            total_score += 0.10
        elif headings_ok > 0:
            print(
                f"FAIL: Component 5 — Only {headings_ok}/{heading_total} headings are 14pt bold"
            )
        else:
            print(
                f"FAIL: Component 5 — No article headings found at 14pt bold "
                f"(checked {heading_total} paragraphs)"
            )
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM environment
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
