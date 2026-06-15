"""
Initial Setup: Q1 Client Newsletter - pre-task state
Task ID: writer_mktg_006
Domain: libreoffice_writer

Creates q1_client_newsletter.docx with:
- Single-column layout
- 1-inch margins on all sides
- No column separator
- Newsletter title, date, and 6 article sections
- Article headings NOT 14pt bold (initial state)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'q1_client_newsletter'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set 1-inch margins on all sides (initial state)
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    # Single column (default — no multi-column XML)

    # --- Newsletter Header (Title + Date) ---
    # Title paragraph
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('The Vanguard Quarterly \u2014 Q1 2026')
    title_run.font.name = 'Georgia'
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(6)

    # Date line
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_para.add_run('January 2026 \u2022 Quarterly Client Edition')
    date_run.font.name = 'Georgia'
    date_run.font.size = Pt(11)
    date_run.italic = True
    date_para.paragraph_format.space_before = Pt(0)
    date_para.paragraph_format.space_after = Pt(18)

    # Horizontal rule placeholder (empty paragraph with border would be complex; use a paragraph break)
    separator = doc.add_paragraph()
    separator.paragraph_format.space_before = Pt(0)
    separator.paragraph_format.space_after = Pt(12)

    # --- Article 1: Market Outlook ---
    h1 = doc.add_paragraph()
    h1_run = h1.add_run('Market Outlook: Navigating Uncertainty in 2026')
    h1_run.font.name = 'Calibri'
    h1_run.font.size = Pt(12)
    h1_run.bold = True

    p1 = doc.add_paragraph()
    p1_run = p1.add_run(
        'As we enter the first quarter of 2026, global markets continue to recalibrate '
        'following last year\u2019s interest rate adjustments. Equity markets have shown '
        'resilience, with the S&P 500 posting modest gains of 3.2% through December. '
        'Our analysts anticipate continued volatility in the near term as central banks '
        'navigate the final mile of inflation control.'
    )
    p1_run.font.name = 'Calibri'
    p1_run.font.size = Pt(10)
    p1.paragraph_format.space_after = Pt(8)

    p1b = doc.add_paragraph()
    p1b_run = p1b.add_run(
        'Key sectors to watch include healthcare technology, renewable energy infrastructure, '
        'and emerging market bonds. Our portfolio managers have positioned client holdings '
        'defensively while maintaining exposure to high-conviction growth opportunities.'
    )
    p1b_run.font.name = 'Calibri'
    p1b_run.font.size = Pt(10)
    p1b.paragraph_format.space_after = Pt(12)

    # --- Article 2: Portfolio Performance ---
    h2 = doc.add_paragraph()
    h2_run = h2.add_run('Portfolio Performance: Q4 2025 Review')
    h2_run.font.name = 'Calibri'
    h2_run.font.size = Pt(12)
    h2_run.bold = True

    p2 = doc.add_paragraph()
    p2_run = p2.add_run(
        'The fourth quarter of 2025 delivered strong results across our balanced and '
        'growth portfolio strategies. The Vanguard Growth Portfolio returned 6.8% for '
        'the quarter, outperforming its benchmark by 210 basis points. Fixed income '
        'allocations provided crucial stability as equity markets experienced mid-quarter '
        'turbulence in October.'
    )
    p2_run.font.name = 'Calibri'
    p2_run.font.size = Pt(10)
    p2.paragraph_format.space_after = Pt(8)

    p2b = doc.add_paragraph()
    p2b_run = p2b.add_run(
        'Technology holdings — particularly positions in semiconductor and AI infrastructure '
        'companies — were the primary drivers of outperformance. We trimmed these positions '
        'in November to lock in gains and rebalance sector allocations back to target weights.'
    )
    p2b_run.font.name = 'Calibri'
    p2b_run.font.size = Pt(10)
    p2b.paragraph_format.space_after = Pt(12)

    # --- Article 3: Tax Planning Strategies ---
    h3 = doc.add_paragraph()
    h3_run = h3.add_run('Tax Planning: Strategies for the New Fiscal Year')
    h3_run.font.name = 'Calibri'
    h3_run.font.size = Pt(12)
    h3_run.bold = True

    p3 = doc.add_paragraph()
    p3_run = p3.add_run(
        'With the tax year now underway, this is an ideal time to review your tax-loss '
        'harvesting opportunities and Roth conversion strategies. The 2026 contribution '
        'limits for 401(k) plans have increased to $23,500, and IRA limits remain at '
        '$7,000 ($8,000 for those aged 50 and over). We encourage all clients to maximize '
        'these tax-advantaged vehicles early in the year.'
    )
    p3_run.font.name = 'Calibri'
    p3_run.font.size = Pt(10)
    p3.paragraph_format.space_after = Pt(8)

    p3b = doc.add_paragraph()
    p3b_run = p3b.add_run(
        'Estate planning clients should note that the federal estate tax exemption is '
        'scheduled for legislative review this year. Our estate planning team is monitoring '
        'developments closely and will reach out individually to affected clients with '
        'recommended action steps.'
    )
    p3b_run.font.name = 'Calibri'
    p3b_run.font.size = Pt(10)
    p3b.paragraph_format.space_after = Pt(12)

    # --- Article 4: Real Estate Update ---
    h4 = doc.add_paragraph()
    h4_run = h4.add_run('Real Estate Investments: Trends and Opportunities')
    h4_run.font.name = 'Calibri'
    h4_run.font.size = Pt(12)
    h4_run.bold = True

    p4 = doc.add_paragraph()
    p4_run = p4.add_run(
        'Commercial real estate continues its uneven recovery, with industrial and '
        'logistics properties outperforming traditional office assets. Our REIT allocations '
        'have been shifted toward data center operators and cell tower companies, which '
        'benefit from the accelerating buildout of AI compute infrastructure. Residential '
        'real estate in Sun Belt metropolitan areas remains a compelling long-term hold.'
    )
    p4_run.font.name = 'Calibri'
    p4_run.font.size = Pt(10)
    p4.paragraph_format.space_after = Pt(8)

    p4b = doc.add_paragraph()
    p4b_run = p4b.add_run(
        'Mortgage rates have eased slightly from their 2024 highs, though they remain '
        'elevated relative to the 2021\u20132022 period. Clients considering property '
        'purchases should consult with our lending partners, who can offer competitive '
        'terms through our preferred relationship programs.'
    )
    p4b_run.font.name = 'Calibri'
    p4b_run.font.size = Pt(10)
    p4b.paragraph_format.space_after = Pt(12)

    # --- Article 5: ESG Investing ---
    h5 = doc.add_paragraph()
    h5_run = h5.add_run('ESG Investing: Balancing Values and Returns')
    h5_run.font.name = 'Calibri'
    h5_run.font.size = Pt(12)
    h5_run.bold = True

    p5 = doc.add_paragraph()
    p5_run = p5.add_run(
        'Environmental, social, and governance-focused portfolios delivered competitive '
        'returns in 2025, with our ESG Growth strategy returning 7.4% for the year — '
        'slightly ahead of its traditional counterpart. Increased regulatory pressure '
        'on carbon disclosure has begun to differentiate companies with genuine sustainability '
        'practices from those engaged in greenwashing.'
    )
    p5_run.font.name = 'Calibri'
    p5_run.font.size = Pt(10)
    p5.paragraph_format.space_after = Pt(8)

    p5b = doc.add_paragraph()
    p5b_run = p5b.add_run(
        'We have expanded our ESG screening criteria to include supply chain labor '
        'practices and water resource management metrics. Clients interested in aligning '
        'their portfolios with their values are encouraged to schedule a consultation '
        'with their relationship manager to discuss our full range of impact investment options.'
    )
    p5b_run.font.name = 'Calibri'
    p5b_run.font.size = Pt(10)
    p5b.paragraph_format.space_after = Pt(12)

    # --- Article 6: Client Services Update ---
    h6 = doc.add_paragraph()
    h6_run = h6.add_run('Client Services: New Features in Your Online Portal')
    h6_run.font.name = 'Calibri'
    h6_run.font.size = Pt(12)
    h6_run.bold = True

    p6 = doc.add_paragraph()
    p6_run = p6.add_run(
        'We are pleased to announce several enhancements to the Vanguard client portal, '
        'rolling out in Q1 2026. The updated dashboard now provides real-time allocation '
        'tracking, customizable performance benchmarks, and integrated tax document access. '
        'Mobile app users will benefit from improved biometric authentication and '
        'streamlined transaction workflows.'
    )
    p6_run.font.name = 'Calibri'
    p6_run.font.size = Pt(10)
    p6.paragraph_format.space_after = Pt(8)

    p6b = doc.add_paragraph()
    p6b_run = p6b.add_run(
        'Our dedicated client service team is available Monday through Friday, 8 AM to '
        '6 PM Eastern Time, to assist with portal navigation and account inquiries. '
        'You can also reach us through the secure messaging feature within the portal '
        'for non-urgent requests. We look forward to serving you throughout this new year.'
    )
    p6b_run.font.name = 'Calibri'
    p6b_run.font.size = Pt(10)
    p6b.paragraph_format.space_after = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
