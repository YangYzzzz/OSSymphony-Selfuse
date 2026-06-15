"""
Initial Setup: Create a Writer document with 8 technical terms and definitions as plain text.
Task ID: writer_lec_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_031'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)

# Glossary terms and definitions (technical computing terms)
GLOSSARY = [
    (
        "API (Application Programming Interface)",
        "A set of protocols, routines, and tools that specifies how software components should interact with each other. APIs define the methods and data formats that applications use to communicate, enabling developers to integrate external services without understanding their internal implementation."
    ),
    (
        "Latency",
        "The time delay between the initiation of a request and the receipt of a response in a computing system. Network latency is typically measured in milliseconds and is influenced by factors such as physical distance, network congestion, and the number of routing hops between endpoints."
    ),
    (
        "Containerization",
        "A lightweight virtualization method that packages an application together with its dependencies, libraries, and configuration files into a single portable unit called a container. Unlike traditional virtual machines, containers share the host operating system kernel, resulting in faster startup times and lower resource overhead."
    ),
    (
        "Idempotency",
        "A property of certain operations in mathematics and computer science where performing the same operation multiple times produces the same result as performing it once. In RESTful API design, idempotent methods such as GET, PUT, and DELETE guarantee that repeated identical requests will not cause unintended side effects."
    ),
    (
        "Sharding",
        "A database architecture pattern in which data is horizontally partitioned across multiple independent database instances called shards. Each shard holds a subset of the total data, allowing the system to distribute read and write operations and scale beyond the capacity of a single server."
    ),
    (
        "Eventual Consistency",
        "A consistency model used in distributed computing that guarantees all replicas of a data item will converge to the same value given sufficient time without new updates. This model trades immediate consistency for improved availability and partition tolerance, as described by the CAP theorem."
    ),
    (
        "Garbage Collection",
        "An automatic memory management mechanism that identifies and reclaims memory occupied by objects that are no longer reachable or referenced by the running program. Modern garbage collectors employ generational algorithms that segregate objects by age, collecting short-lived objects more frequently to optimize throughput and minimize pause times."
    ),
    (
        "Load Balancing",
        "The process of distributing incoming network traffic or computational workloads across multiple servers or resources to ensure no single server becomes overwhelmed. Common algorithms include round-robin, least connections, and weighted distribution, each optimizing for different performance characteristics."
    ),
]

def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add a title
    title = doc.add_heading('Technical Glossary', level=1)

    # Add each term and definition as plain paragraphs (no special formatting)
    for term, definition in GLOSSARY:
        # Term as a plain paragraph
        term_para = doc.add_paragraph(term)
        # Definition as a plain paragraph
        def_para = doc.add_paragraph(definition)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')

create_initial()
