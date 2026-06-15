"""
Initial Setup: Create a Writer document with realistic content but no custom styles.
Task ID: writer_bs_067
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_067'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Quarterly Operations Report — Q1 2025', level=0)

    # --- Section 1: Executive Summary ---
    doc.add_heading('Executive Summary', level=1)

    p1 = doc.add_paragraph(
        'The first quarter of 2025 has demonstrated significant progress across '
        'all operational divisions. Revenue targets were exceeded by 8.3%, driven '
        'primarily by the expansion of our digital services portfolio. Customer '
        'acquisition costs decreased by 12% compared to Q4 2024, while retention '
        'rates improved to 94.7%.'
    )

    p2 = doc.add_paragraph(
        'Infrastructure investments completed during this period include the '
        'deployment of three new regional data centers in Frankfurt, Singapore, '
        'and São Paulo. These facilities are expected to reduce latency by 40% '
        'for users in EMEA and APAC regions.'
    )

    # --- Section 2: Financial Overview ---
    doc.add_heading('Financial Overview', level=1)

    p3 = doc.add_paragraph(
        'Total revenue for Q1 2025 reached $48.7 million, representing a 15.2% '
        'year-over-year increase. The breakdown by division is as follows:'
    )

    # Add a simple table for financial data
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Division', 'Revenue (USD)', 'YoY Growth']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Enterprise Solutions', '$18,240,000', '+12.4%'],
        ['Digital Services', '$15,890,000', '+23.1%'],
        ['Consulting', '$9,320,000', '+8.7%'],
        ['Support & Maintenance', '$5,250,000', '+4.2%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()  # spacing

    # --- Section 3: Operational Highlights ---
    doc.add_heading('Operational Highlights', level=1)

    p4 = doc.add_paragraph(
        'Several key milestones were achieved during this quarter. The engineering '
        'team successfully migrated 78% of legacy applications to the new '
        'microservices architecture, ahead of the projected 65% target. This '
        'migration has already yielded measurable improvements in system '
        'reliability, with unplanned downtime reduced by 62%.'
    )

    # A quote-like paragraph (but using default style, no custom style)
    p5 = doc.add_paragraph(
        '"The transition to cloud-native infrastructure represents the most '
        'significant technological shift in our company\'s history. We are not '
        'merely upgrading systems — we are fundamentally reimagining how we '
        'deliver value to our customers." — Elena Vasquez, CTO'
    )

    p6 = doc.add_paragraph(
        'The customer success team expanded its coverage to include dedicated '
        'account managers for all enterprise clients generating over $500,000 '
        'in annual revenue. This initiative has already shown positive results, '
        'with enterprise NPS scores increasing from 72 to 81.'
    )

    # --- Section 4: Risk Assessment ---
    doc.add_heading('Risk Assessment', level=1)

    p7 = doc.add_paragraph(
        'While overall performance has been strong, several risk factors require '
        'ongoing monitoring. Supply chain disruptions in the semiconductor sector '
        'continue to affect hardware procurement timelines. The average lead time '
        'for server equipment has increased from 8 weeks to 14 weeks.'
    )

    # A note-like paragraph (but using default style, no custom style)
    p8 = doc.add_paragraph(
        'Note: All financial figures are preliminary and subject to final audit '
        'confirmation. Adjusted figures will be published in the annual report. '
        'Currency conversions use the Q1 2025 average exchange rate of '
        'EUR/USD 1.0842.'
    )

    # --- Section 5: Strategic Outlook ---
    doc.add_heading('Strategic Outlook', level=1)

    p9 = doc.add_paragraph(
        'Looking ahead to Q2 2025, the leadership team has identified three '
        'priority areas for investment: artificial intelligence integration '
        'across the product suite, expansion into the Latin American market, '
        'and the launch of a new developer platform. Initial budget allocations '
        'total $12.5 million for these initiatives.'
    )

    p10 = doc.add_paragraph(
        'The Board of Directors has approved a revised compensation framework '
        'designed to attract and retain top talent in competitive markets. '
        'The new framework includes equity participation for senior individual '
        'contributors and enhanced professional development budgets.'
    )

    # --- Section 6: Conclusion ---
    doc.add_heading('Conclusion', level=1)

    p11 = doc.add_paragraph(
        'Q1 2025 results confirm that the strategic direction established in '
        'the 2024 annual planning cycle is delivering measurable outcomes. '
        'The organization remains well-positioned to capitalize on market '
        'opportunities while maintaining operational discipline. Detailed '
        'departmental reports are available in the appendices.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
