"""
Initial Setup: color_grid.docx with 4 scattered rectangle shapes
Task ID: writer_obj_058
Domain: libreoffice_writer

Creates a .docx file with 4 colored rectangles at random positions/sizes.
The agent needs to arrange them into a 2x2 grid.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_058'
OUTPUT = f'{WORKDIR}/Desktop/color_grid.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def cm_to_emu(cm_val):
    """Convert centimeters to English Metric Units (EMU)."""
    return int(cm_val * 360000)


def make_rect_shape(shape_id, name, fill_color_hex, x_cm, y_cm, w_cm, h_cm):
    """
    Create a VML/DrawingML rectangle shape as an XML element for insertion into docx.
    We use the drawing XML (wp:anchor + a:graphic) approach with DML.
    """
    # Convert to EMU
    x_emu = cm_to_emu(x_cm)
    y_emu = cm_to_emu(y_cm)
    w_emu = cm_to_emu(w_cm)
    h_emu = cm_to_emu(h_cm)

    # Parse hex color
    r = int(fill_color_hex[1:3], 16)
    g = int(fill_color_hex[3:5], 16)
    b = int(fill_color_hex[5:7], 16)
    color_hex = f'{r:02X}{g:02X}{b:02X}'

    # Build the XML for an anchored drawing with a preset geometry rectangle
    xml_str = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
               xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:r>
    <w:rPr/>
    <w:drawing>
      <wp:anchor distT="0" distB="0" distL="0" distR="0"
                 simplePos="0" relativeHeight="251659264" behindDoc="0"
                 locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page">
          <wp:posOffset>{x_emu}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="page">
          <wp:posOffset>{y_emu}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{w_emu}" cy="{h_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="{shape_id}" name="{name}"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
            <pic:pic>
              <pic:nvPicPr>
                <pic:cNvPr id="{shape_id}" name="{name}"/>
                <pic:cNvPicPr preferRelativeResize="0">
                  <a:picLocks noChangeAspect="0"/>
                </pic:cNvPicPr>
              </pic:nvPicPr>
              <pic:blipFill>
                <a:stretch><a:fillRect/></a:stretch>
              </pic:blipFill>
              <pic:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{w_emu}" cy="{h_emu}"/>
                </a:xfrm>
                <a:prstGeom prst="rect">
                  <a:avLst/>
                </a:prstGeom>
                <a:solidFill>
                  <a:srgbClr val="{color_hex}"/>
                </a:solidFill>
                <a:ln w="9525">
                  <a:solidFill>
                    <a:srgbClr val="000000"/>
                  </a:solidFill>
                </a:ln>
              </pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </w:r>
</w:p>'''
    return etree.fromstring(xml_str)


def make_sp_shape(shape_id, name, fill_color_hex, x_cm, y_cm, w_cm, h_cm):
    """
    Create a proper sp (shape) drawing element using wps namespace.
    This creates a proper rectangle shape in docx.
    """
    x_emu = cm_to_emu(x_cm)
    y_emu = cm_to_emu(y_cm)
    w_emu = cm_to_emu(w_cm)
    h_emu = cm_to_emu(h_cm)

    # Parse hex color
    color_hex = fill_color_hex.lstrip('#').upper()

    xml_str = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
               xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
  <w:r>
    <w:rPr/>
    <w:drawing>
      <wp:anchor distT="0" distB="0" distL="114300" distR="114300"
                 simplePos="0" relativeHeight="251659264" behindDoc="0"
                 locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page">
          <wp:posOffset>{x_emu}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="page">
          <wp:posOffset>{y_emu}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{w_emu}" cy="{h_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="{shape_id}" name="{name}" descr="{name}"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp>
              <wps:cNvSpPr>
                <a:spLocks noChangeArrowheads="1"/>
              </wps:cNvSpPr>
              <wps:spPr bwMode="auto">
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{w_emu}" cy="{h_emu}"/>
                </a:xfrm>
                <a:prstGeom prst="rect">
                  <a:avLst/>
                </a:prstGeom>
                <a:solidFill>
                  <a:srgbClr val="{color_hex}"/>
                </a:solidFill>
                <a:ln w="12700">
                  <a:solidFill>
                    <a:srgbClr val="000000"/>
                  </a:solidFill>
                </a:ln>
              </wps:spPr>
              <wps:bodyPr/>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </w:r>
</w:p>'''
    return etree.fromstring(xml_str)


def create_initial():
    # Ensure Desktop directory exists
    desktop_dir = f'{WORKDIR}/Desktop'
    os.makedirs(desktop_dir, exist_ok=True)

    doc = Document()

    # Set standard A4 page dimensions
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Add a title paragraph
    title_para = doc.add_paragraph()
    run = title_para.add_run("Color Grid Layout")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph("Arrange the four shapes into a 2x2 grid.")

    # Add shapes as drawing elements scattered at random positions/sizes
    # Shapes are randomly positioned and with different sizes (not in a grid)
    # Red rectangle - top area, random size
    # Shape positions (scattered, not in 2x2 grid):
    # Red: X=3cm, Y=4cm, W=8cm, H=3cm (not 6x4)
    # Blue: X=11cm, Y=3cm, W=5cm, H=5cm (not 6x4)
    # Green: X=4cm, Y=12cm, W=7cm, H=2.5cm (not 6x4)
    # Yellow: X=10cm, Y=14cm, W=4cm, H=6cm (not 6x4)

    shapes_data = [
        # (shape_id, name, color_hex, x_cm, y_cm, w_cm, h_cm)
        (1, 'Red Rectangle', '#F44336', 3.0, 4.0, 8.0, 3.0),
        (2, 'Blue Rectangle', '#1565C0', 11.0, 3.0, 5.0, 5.0),
        (3, 'Green Rectangle', '#4CAF50', 4.0, 12.0, 7.0, 2.5),
        (4, 'Yellow Rectangle', '#FFC107', 10.0, 14.0, 4.0, 6.0),
    ]

    body = doc.element.body
    for shape_id, name, color, x, y, w, h in shapes_data:
        shape_elem = make_sp_shape(shape_id, name, color, x, y, w, h)
        body.append(shape_elem)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
