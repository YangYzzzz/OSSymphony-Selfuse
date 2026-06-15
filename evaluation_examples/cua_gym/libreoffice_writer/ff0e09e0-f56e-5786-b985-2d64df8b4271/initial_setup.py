"""
Initial Setup: Speech translation document with blank paragraphs between content paragraphs
Task ID: wrpara_047
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'wrpara_047'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# The 12 content paragraphs of a translated speech
PARAGRAPHS = [
    "Distinguished delegates, honored guests, and fellow citizens of the world — it is my profound privilege to address this assembly today on a matter that transcends borders, ideologies, and generations.",
    "We gather at a pivotal moment in history. The challenges before us — climate instability, economic inequality, and the erosion of trust in institutions — are not confined to any single nation. They are shared burdens that demand collective resolve and unprecedented cooperation.",
    "Let me begin with what we know. Over the past decade, global temperatures have risen by an average of 0.3 degrees Celsius. Coastal communities from Bangladesh to the Bahamas face existential threats. Agricultural yields in sub-Saharan Africa have declined by 12 percent, displacing millions of families who once relied on the land their ancestors cultivated for centuries.",
    "But knowledge without action is merely an exercise in documentation. Our grandchildren will not judge us by the reports we published or the conferences we attended. They will judge us by the forests still standing, the rivers still flowing, and the air still breathable.",
    "That is why I am announcing today the Green Horizon Initiative — a multi-lateral framework committing 47 nations to reduce carbon emissions by 40 percent before 2035. This is not aspirational rhetoric. Each participating country has submitted binding implementation plans reviewed by independent auditors.",
    "The initiative rests on three pillars. First, a transition to renewable energy infrastructure funded by a cooperative investment pool of 280 billion dollars. Second, reforestation targets covering 15 million hectares of degraded land across four continents. Third, a technology-sharing agreement ensuring that innovations in carbon capture and sustainable agriculture are accessible to developing nations without prohibitive licensing costs.",
    "Critics will say this is too ambitious. I respectfully disagree. What is truly ambitious — dangerously so — is the belief that we can continue on our current trajectory without catastrophic consequences. The cost of inaction dwarfs the investment required for transformation.",
    "I also want to address the economic dimension. Sustainable development is not the enemy of prosperity; it is its most reliable engine. Countries that invested early in clean energy — Denmark, Costa Rica, South Korea — have seen consistent GDP growth alongside declining emissions. The notion that environmental responsibility demands economic sacrifice is a false binary we must finally discard.",
    "Furthermore, we must ensure that this transition is just. Workers in fossil fuel industries deserve retraining programs, pension protections, and pathways to new livelihoods. Communities built around coal mines and oil refineries must not be abandoned. The Green Horizon Initiative allocates 18 billion dollars specifically for workforce transition and community resilience.",
    "To the young people listening today — and I know many of you are — I want to say this directly: your frustration is valid. Your impatience is warranted. But I ask you to channel that energy not only into protest but also into participation. Run for office. Start companies. Teach in classrooms. Build the institutions you wish existed. The future is not something that happens to you; it is something you construct.",
    "In closing, let me share a proverb from the Maasai people of East Africa: 'We do not inherit the earth from our ancestors; we borrow it from our children.' Every decision we make in this chamber, every policy we enact, every investment we authorize must be weighed against that simple truth.",
    "The path forward will not be easy. There will be setbacks, disagreements, and moments of doubt. But if we choose courage over convenience, collaboration over competition, and legacy over short-term gain, I believe we will look back on this moment as the turning point. Thank you."
]


def create_initial():
    doc = Document()

    # Set default style to have no extra spacing (standard Normal style)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    for i, text in enumerate(PARAGRAPHS):
        # Add content paragraph
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)

        # Add blank paragraph after each content paragraph (except the last)
        if i < len(PARAGRAPHS) - 1:
            blank = doc.add_paragraph('')
            blank.paragraph_format.space_after = Pt(0)
            blank.paragraph_format.space_before = Pt(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')
    print(f'Total paragraphs: {len(doc.paragraphs)} (12 content + 11 blank = 23)')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
