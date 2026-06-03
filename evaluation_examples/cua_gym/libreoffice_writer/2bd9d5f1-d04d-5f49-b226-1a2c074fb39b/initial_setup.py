"""
Initial Setup: Employee Recognition Program - rough draft with bullet-point notes
Task ID: writer_hr_089
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_089'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Employee_Recognition_Program', level=0)

    # Simple intro paragraph
    doc.add_paragraph(
        'This document contains planning notes for the upcoming employee recognition program. '
        'Each section below outlines key ideas and considerations that need to be developed '
        'into a polished program document.'
    )

    # --- Section 1: Program Philosophy Notes ---
    doc.add_heading('Program Philosophy (Draft Notes)', level=1)
    doc.add_paragraph(
        'We believe that recognizing employees is essential for maintaining a motivated, '
        'productive, and engaged workforce. The following points should be incorporated '
        'into the final philosophy statement:'
    )
    doc.add_paragraph('Timely recognition reinforces positive behaviors', style='List Bullet')
    doc.add_paragraph('Recognition should be accessible to all levels of the organization', style='List Bullet')
    doc.add_paragraph('Both peer-to-peer and manager-driven recognition are valuable', style='List Bullet')
    doc.add_paragraph('Celebrate achievements publicly to inspire others', style='List Bullet')
    doc.add_paragraph('Link recognition to company core values and strategic goals', style='List Bullet')

    # --- Section 2: Recognition Tiers Notes ---
    doc.add_heading('Recognition Tiers (Draft Notes)', level=1)
    doc.add_paragraph('Need to finalize the following tiers and create a detailed table:')
    doc.add_paragraph('Spot Award - small immediate recognition, around $25-50 range, for day-to-day excellence', style='List Bullet')
    doc.add_paragraph('Monthly Star - $100 value, selected each month by department managers', style='List Bullet')
    doc.add_paragraph('Quarterly Champion - $250 value, cross-departmental nomination process', style='List Bullet')
    doc.add_paragraph('Annual MVP - $1000 value, highest honor, voted by leadership team', style='List Bullet')
    doc.add_paragraph('Each tier needs: clear criteria, description, award amount, nomination process', style='List Bullet')

    # --- Section 3: Nomination Process Notes ---
    doc.add_heading('Nomination Form & Scoring (Draft Notes)', level=1)
    doc.add_paragraph('The nomination form should capture:')
    doc.add_paragraph('Nominee name, department, position', style='List Bullet')
    doc.add_paragraph('Nominator name and relationship to nominee', style='List Bullet')
    doc.add_paragraph('Category of recognition (which tier)', style='List Bullet')
    doc.add_paragraph('Description of achievement or behavior', style='List Bullet')
    doc.add_paragraph('Impact on team/organization', style='List Bullet')
    doc.add_paragraph('Alignment with company values', style='List Bullet')
    doc.add_paragraph('')
    doc.add_paragraph('Scoring rubric ideas:')
    doc.add_paragraph('Impact (1-5 scale) - measurable outcomes', style='List Bullet')
    doc.add_paragraph('Innovation (1-5 scale) - creative problem solving', style='List Bullet')
    doc.add_paragraph('Teamwork (1-5 scale) - collaboration and support', style='List Bullet')
    doc.add_paragraph('Values Alignment (1-5 scale) - reflects company mission', style='List Bullet')
    doc.add_paragraph('Consistency (1-5 scale) - sustained excellence', style='List Bullet')

    # --- Section 4: Past Winners Notes ---
    doc.add_heading('Past Winners Showcase (Draft Notes)', level=1)
    doc.add_paragraph('Include a showcase of previous award recipients to inspire nominations:')
    doc.add_paragraph('Q1 2025 - Elena Rodriguez, Engineering, Quarterly Champion - Led migration to cloud platform', style='List Bullet')
    doc.add_paragraph('Q2 2025 - David Kim, Customer Success, Monthly Star - 98% satisfaction rating', style='List Bullet')
    doc.add_paragraph('Q3 2025 - Priya Sharma, Marketing, Annual MVP - Rebranding campaign increased leads 40%', style='List Bullet')
    doc.add_paragraph('Q4 2025 - James O\'Brien, Operations, Spot Award - Streamlined onboarding process', style='List Bullet')
    doc.add_paragraph('Q1 2026 - Maria Santos, Finance, Quarterly Champion - Automated reporting workflow', style='List Bullet')

    # --- Section 5: Budget Notes ---
    doc.add_heading('Budget Allocation (Draft Notes)', level=1)
    doc.add_paragraph('Need to create a department budget table. Preliminary allocations:')
    doc.add_paragraph('Engineering - largest department, highest allocation (~$8,000)', style='List Bullet')
    doc.add_paragraph('Sales - strong performance incentives needed (~$6,500)', style='List Bullet')
    doc.add_paragraph('Marketing - creative recognition opportunities (~$5,000)', style='List Bullet')
    doc.add_paragraph('Customer Success - frontline recognition (~$4,500)', style='List Bullet')
    doc.add_paragraph('Finance - operational excellence focus (~$3,500)', style='List Bullet')
    doc.add_paragraph('Human Resources - leading by example (~$3,000)', style='List Bullet')
    doc.add_paragraph('Operations - efficiency and reliability (~$4,000)', style='List Bullet')
    doc.add_paragraph('Total budget estimate: approximately $34,500', style='List Bullet')

    # --- Section 6: Calendar Notes ---
    doc.add_heading('Program Calendar (Draft Notes)', level=1)
    doc.add_paragraph('Monthly themes to keep the program fresh and engaging:')
    doc.add_paragraph('January - New Year, New Goals: Goal Setting Excellence', style='List Bullet')
    doc.add_paragraph('February - Heart of the Team: Collaboration Award', style='List Bullet')
    doc.add_paragraph('March - Women in Leadership: Mentorship Recognition', style='List Bullet')
    doc.add_paragraph('April - Innovation Month: Creative Solutions Award', style='List Bullet')
    doc.add_paragraph('May - Customer First: Service Excellence', style='List Bullet')
    doc.add_paragraph('June - Mid-Year Milestone: Progress Champion', style='List Bullet')
    doc.add_paragraph('July - Summer of Learning: Professional Growth', style='List Bullet')
    doc.add_paragraph('August - Back to Basics: Operational Excellence', style='List Bullet')
    doc.add_paragraph('September - Team Spirit: Cross-Department Collaboration', style='List Bullet')
    doc.add_paragraph('October - Safety & Wellness: Well-being Champion', style='List Bullet')
    doc.add_paragraph('November - Gratitude Month: Peer Appreciation', style='List Bullet')
    doc.add_paragraph('December - Year in Review: Annual MVP Ceremony', style='List Bullet')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
