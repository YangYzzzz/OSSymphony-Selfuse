"""
Reward Script: Employee Recognition Program Document
Task ID: writer_hr_089
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20) - Recognition Tiers table with 4 tiers
  Component 2 (0.15) - Scoring Rubric table with 5 criteria
  Component 3 (0.15) - Past Winners table with entries
  Component 4 (0.15) - Budget Allocation table with >=6 departments
  Component 5 (0.15) - Program Calendar table with 12 months
  Component 6 (0.20) - Gold/dark blue color scheme on headings and table headers
"""

import os
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_089'


def persist_app_state(domain):
    """Save any open LibreOffice document before verification."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        import time
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def get_cell_shading(cell):
    """Extract cell background fill color from XML."""
    tc = cell._tc
    tc_pr = tc.find(qn('w:tcPr'))
    if tc_pr is not None:
        shd = tc_pr.find(qn('w:shd'))
        if shd is not None:
            return shd.get(qn('w:fill'))
    return None


def color_close(c1_hex, c2_hex, threshold=60):
    """Check if two hex colors are close enough (RGB Euclidean distance)."""
    try:
        r1, g1, b1 = int(c1_hex[0:2], 16), int(c1_hex[2:4], 16), int(c1_hex[4:6], 16)
        r2, g2, b2 = int(c2_hex[0:2], 16), int(c2_hex[2:4], 16), int(c2_hex[4:6], 16)
        dist = ((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2) ** 0.5
        return dist < threshold
    except Exception:
        return False


def find_table_by_header(tables, header_keywords):
    """Find a table whose first row contains cells matching header keywords."""
    for t in tables:
        if len(t.rows) < 1:
            continue
        row0_texts = [c.text.strip().lower() for c in t.rows[0].cells]
        if all(any(kw in ct for ct in row0_texts) for kw in header_keywords):
            return t
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    tables = doc.tables

    # Component 1: Recognition Tiers table with 4 tiers (0.20 points)
    # Golden has 5 rows x 4 cols: header + Spot Award, Monthly Star, Quarterly Champion, Annual MVP
    try:
        tiers_table = find_table_by_header(tables, ['award', 'value'])
        if tiers_table is None:
            tiers_table = find_table_by_header(tables, ['tier', 'description'])

        if tiers_table is not None and len(tiers_table.rows) >= 5:
            tier_names = []
            for row in list(tiers_table.rows)[1:]:
                cell_text = row.cells[0].text.strip().lower()
                tier_names.append(cell_text)

            expected_tiers = ['spot award', 'monthly star', 'quarterly champion', 'annual mvp']
            matched = sum(1 for et in expected_tiers if any(et in tn for tn in tier_names))

            if matched == 4:
                print(f"PASS: Component 1 - Recognition Tiers table with all 4 tiers found (0.20 pts)")
                total_score += 0.20
            elif matched >= 2:
                partial = 0.10
                print(f"PARTIAL: Component 1 - Recognition Tiers table found with {matched}/4 tiers ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 1 - Tiers table found but only {matched}/4 tiers matched")
        else:
            print(f"FAIL: Component 1 - Recognition Tiers table not found or too few rows (tables={len(tables)})")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Scoring Rubric table with 5 criteria (0.15 points)
    # Golden has 6 rows x 4 cols: header + Impact, Innovation, Teamwork, Values Alignment, Consistency
    try:
        rubric_table = find_table_by_header(tables, ['criteria', 'weight'])
        if rubric_table is None:
            rubric_table = find_table_by_header(tables, ['criteria', 'score'])

        if rubric_table is not None and len(rubric_table.rows) >= 6:
            criteria_names = []
            for row in list(rubric_table.rows)[1:]:
                criteria_names.append(row.cells[0].text.strip().lower())

            expected_criteria = ['impact', 'innovation', 'teamwork', 'values', 'consistency']
            matched = sum(1 for ec in expected_criteria if any(ec in cn for cn in criteria_names))

            if matched >= 4:
                print(f"PASS: Component 2 - Scoring Rubric table with {matched}/5 criteria (0.15 pts)")
                total_score += 0.15
            elif matched >= 2:
                partial = 0.07
                print(f"PARTIAL: Component 2 - Rubric table found with {matched}/5 criteria ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 - Rubric table found but only {matched}/5 criteria matched")
        else:
            print(f"FAIL: Component 2 - Scoring Rubric table not found or too few rows")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Past Winners table with entries (0.15 points)
    # Golden has 8 rows x 5 cols: header + 7 winner entries
    try:
        winners_table = find_table_by_header(tables, ['recipient', 'department'])
        if winners_table is None:
            winners_table = find_table_by_header(tables, ['winner', 'award'])
        if winners_table is None:
            winners_table = find_table_by_header(tables, ['period', 'achievement'])

        if winners_table is not None:
            data_rows = len(winners_table.rows) - 1  # exclude header
            if data_rows >= 5:
                print(f"PASS: Component 3 - Past Winners table with {data_rows} entries (0.15 pts)")
                total_score += 0.15
            elif data_rows >= 3:
                partial = 0.07
                print(f"PARTIAL: Component 3 - Past Winners table with {data_rows} entries ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 - Past Winners table has only {data_rows} entries")
        else:
            print(f"FAIL: Component 3 - Past Winners table not found")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Budget Allocation table with >=6 departments (0.15 points)
    # Golden has 9 rows x 5 cols: header + 7 departments + total row
    try:
        budget_table = find_table_by_header(tables, ['department', 'budget'])
        if budget_table is None:
            budget_table = find_table_by_header(tables, ['department', 'allocation'])

        if budget_table is not None:
            data_rows = len(budget_table.rows) - 1  # exclude header
            # Check for department names
            dept_names = []
            for row in list(budget_table.rows)[1:]:
                dept_text = row.cells[0].text.strip().lower()
                if dept_text and dept_text != 'total':
                    dept_names.append(dept_text)

            if len(dept_names) >= 6:
                print(f"PASS: Component 4 - Budget table with {len(dept_names)} departments (0.15 pts)")
                total_score += 0.15
            elif len(dept_names) >= 4:
                partial = 0.07
                print(f"PARTIAL: Component 4 - Budget table with {len(dept_names)} departments ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 - Budget table has only {len(dept_names)} departments")
        else:
            print(f"FAIL: Component 4 - Budget Allocation table not found")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Program Calendar table with 12 months (0.15 points)
    # Golden has 13 rows x 4 cols: header + 12 months
    try:
        calendar_table = find_table_by_header(tables, ['month', 'theme'])

        if calendar_table is not None:
            month_names = ['january', 'february', 'march', 'april', 'may', 'june',
                           'july', 'august', 'september', 'october', 'november', 'december']
            found_months = []
            for row in list(calendar_table.rows)[1:]:
                month_text = row.cells[0].text.strip().lower()
                for m in month_names:
                    if m in month_text:
                        found_months.append(m)

            if len(found_months) >= 12:
                print(f"PASS: Component 5 - Calendar table with all 12 months (0.15 pts)")
                total_score += 0.15
            elif len(found_months) >= 6:
                partial = 0.07
                print(f"PARTIAL: Component 5 - Calendar table with {len(found_months)}/12 months ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 - Calendar table found but only {len(found_months)} months")
        else:
            print(f"FAIL: Component 5 - Program Calendar table not found")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: Gold/dark blue color scheme (0.20 points)
    # Headings should be dark blue (~003366), table headers should have dark blue shading + gold text (~FFD700)
    try:
        DARK_BLUE = '003366'
        GOLD = 'FFD700'

        # Check 6a: Heading colors are dark blue (0.10 pts)
        dark_blue_headings = 0
        total_headings = 0
        for para in doc.paragraphs:
            if para.style and 'Heading' in para.style.name:
                total_headings += 1
                for run in para.runs:
                    if run.font.color and run.font.color.rgb:
                        rgb_hex = str(run.font.color.rgb)
                        if color_close(rgb_hex, DARK_BLUE):
                            dark_blue_headings += 1
                            break

        heading_pass = total_headings > 0 and dark_blue_headings >= (total_headings * 0.5)

        # Check 6b: Table header rows have dark blue shading and gold text (0.10 pts)
        tables_with_color = 0
        total_tables_checked = min(len(tables), 5)
        for t in tables[:5]:
            if len(t.rows) < 1:
                continue
            # Check header row shading
            dark_blue_cells = 0
            gold_text_cells = 0
            for cell in t.rows[0].cells:
                fill = get_cell_shading(cell)
                if fill and color_close(fill, DARK_BLUE):
                    dark_blue_cells += 1
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.color and run.font.color.rgb:
                            rgb_hex = str(run.font.color.rgb)
                            if color_close(rgb_hex, GOLD):
                                gold_text_cells += 1

            if dark_blue_cells > 0 and gold_text_cells > 0:
                tables_with_color += 1

        table_color_pass = total_tables_checked > 0 and tables_with_color >= (total_tables_checked * 0.5)

        if heading_pass and table_color_pass:
            print(f"PASS: Component 6 - Gold/dark blue color scheme ({dark_blue_headings} headings, {tables_with_color} tables) (0.20 pts)")
            total_score += 0.20
        elif heading_pass or table_color_pass:
            partial = 0.10
            detail = "headings" if heading_pass else "table headers"
            print(f"PARTIAL: Component 6 - Color scheme partially applied ({detail}) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 6 - Gold/dark blue color scheme not detected (headings: {dark_blue_headings}/{total_headings}, tables: {tables_with_color}/{total_tables_checked})")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
