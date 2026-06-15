"""
Initial Setup: Insert 'TODO: Add citation here' after sentence on page 2
Task ID: writer_edit_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_edit_043'
OUTPUT = f'{WORKDIR}/Desktop/literature_review.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page 1 Content: Title and Introduction ---
    title = doc.add_heading('The Future of Work: A Comprehensive Literature Review', level=0)

    intro_heading = doc.add_heading('1. Introduction', level=1)

    intro_para1 = doc.add_paragraph(
        'The nature of work has undergone profound transformations over the past several decades. '
        'Technological advancements, globalization, and shifting organizational priorities have collectively '
        'reshaped how, where, and when people perform their professional duties. This literature review '
        'examines the multifaceted dimensions of contemporary work arrangements, drawing upon empirical '
        'research and theoretical frameworks from organizational behavior, labor economics, and management science.'
    )

    intro_para2 = doc.add_paragraph(
        'The emergence of digital communication technologies has played a pivotal role in enabling flexible '
        'work configurations. From early telecommuting experiments in the 1970s to today\'s sophisticated '
        'remote collaboration platforms, the technological infrastructure supporting distributed work has '
        'grown increasingly robust. Scholars such as Nilles (1975) first conceptualized telecommuting as '
        'a strategy to reduce commuting costs and energy consumption, a vision that has since expanded '
        'dramatically in scope and ambition.'
    )

    intro_para3 = doc.add_paragraph(
        'This review synthesizes findings from peer-reviewed journals, industry reports, and longitudinal '
        'studies published between 2010 and 2024. The analysis focuses on three primary themes: (1) the '
        'drivers and barriers to remote work adoption, (2) the effects of distributed work on individual '
        'and team performance, and (3) the organizational and societal implications of large-scale shifts '
        'in work location and scheduling flexibility.'
    )

    methodology_heading = doc.add_heading('2. Methodology', level=1)

    methodology_para1 = doc.add_paragraph(
        'A systematic literature search was conducted using academic databases including PubMed, JSTOR, '
        'Google Scholar, and the Social Science Research Network. Search terms included "remote work," '
        '"telecommuting," "distributed teams," "flexible work arrangements," and "virtual collaboration." '
        'The initial search yielded over 3,200 potentially relevant articles, which were subsequently '
        'narrowed through abstract screening and full-text review.'
    )

    methodology_para2 = doc.add_paragraph(
        'Inclusion criteria required that studies be published in peer-reviewed venues, include empirical '
        'data collection, and directly address outcomes related to remote or flexible work. Studies '
        'relying solely on anecdotal evidence or practitioner opinion without systematic data collection '
        'were excluded. The final corpus comprised 187 studies spanning 14 countries and representing '
        'over 220,000 individual participants across diverse industries and organizational contexts.'
    )

    # Page break to push content to page 2
    doc.add_page_break()

    # --- Page 2 Content: Remote Work Section ---
    remote_heading = doc.add_heading('3. The Rise of Remote Work', level=1)

    remote_para1 = doc.add_paragraph(
        'Remote work has transitioned from a niche benefit offered by progressive technology companies '
        'to a mainstream employment arrangement embraced across virtually every sector of the global economy. '
        'Early adopters were primarily knowledge workers in software development, consulting, and financial '
        'services, but the phenomenon has since expanded to encompass roles in education, healthcare '
        'administration, customer service, and creative industries.'
    )

    # THE KEY SENTENCE - MUST be here without "TODO: Add citation here"
    remote_para2 = doc.add_paragraph(
        'The prevalence of remote work has increased by 300% since 2020, according to recent studies.'
    )

    remote_para3 = doc.add_paragraph(
        'Several factors have converged to accelerate this trend. Cloud-based productivity suites, '
        'video conferencing platforms, and project management tools have collectively lowered the '
        'technical barriers to remote collaboration. Simultaneously, a generation of digital natives '
        'has entered the workforce with expectations of flexibility and autonomy that traditional '
        'office-centric arrangements may struggle to satisfy.'
    )

    remote_para4 = doc.add_paragraph(
        'Research by Allen, Golden, and Shockley (2015) provides a comprehensive meta-analytic review '
        'of telecommuting outcomes, synthesizing findings from 46 studies conducted between 1980 and 2013. '
        'Their analysis reveals a generally positive relationship between telecommuting and job satisfaction, '
        'though the magnitude of this effect varies substantially across individual characteristics, '
        'job types, and organizational contexts. Workers who reported high levels of professional isolation '
        'demonstrated smaller satisfaction gains from remote arrangements compared to their more socially '
        'independent colleagues.'
    )

    productivity_heading = doc.add_heading('4. Productivity and Performance Outcomes', level=1)

    productivity_para1 = doc.add_paragraph(
        'The relationship between remote work and individual productivity has been the subject of '
        'considerable scholarly debate. Early studies often reported productivity gains among remote workers, '
        'attributing these improvements to reduced commuting fatigue, fewer workplace interruptions, '
        'and greater autonomy over work scheduling. A frequently cited experiment by Bloom et al. (2015) '
        'at a Chinese travel company found a 13% productivity increase among employees assigned to work '
        'from home compared to their office-based counterparts.'
    )

    productivity_para2 = doc.add_paragraph(
        'However, subsequent research has complicated this optimistic narrative. Studies examining '
        'knowledge-intensive collaborative work suggest that the productivity benefits of remote arrangements '
        'may diminish for tasks requiring intensive coordination, real-time problem-solving, or the transfer '
        'of tacit knowledge. Teams engaged in complex innovation projects have shown mixed results, with '
        'some reporting decreased creative output and others maintaining or improving performance through '
        'deliberate virtual collaboration practices.'
    )

    # Save
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
