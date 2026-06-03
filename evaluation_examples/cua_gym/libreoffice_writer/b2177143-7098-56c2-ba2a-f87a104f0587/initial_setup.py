"""
Initial Setup: Financial report document with multiple pages.
Task ID: writer_biz_036
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_036'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_body_para(doc, text, bold_first_sentence=False, space_after=Pt(6)):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = space_after
    if bold_first_sentence and '.' in text:
        idx = text.index('.') + 1
        run_bold = para.add_run(text[:idx])
        run_bold.bold = True
        run_bold.font.size = Pt(11)
        run_bold.font.name = 'Calibri'
        run_rest = para.add_run(text[idx:])
        run_rest.font.size = Pt(11)
        run_rest.font.name = 'Calibri'
    else:
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return para


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # === PAGE 1: Title Page and Executive Summary ===

    # Title
    title = doc.add_heading('Meridian Technologies Inc.', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_heading('Annual Financial Performance Report — FY 2025', level=1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # spacer

    add_body_para(doc, 'Prepared by: Office of the Chief Financial Officer')
    add_body_para(doc, 'Date: March 28, 2025')
    add_body_para(doc, 'Classification: Internal — Confidential')
    add_body_para(doc, 'Distribution: Executive Leadership Team, Board of Directors')

    doc.add_paragraph()  # spacer

    add_heading_styled(doc, 'Executive Summary', level=1)

    add_body_para(doc, (
        'Meridian Technologies Inc. delivered strong financial results in fiscal year 2025, '
        'driven by robust demand for our enterprise cloud platform and sustained growth in the '
        'cybersecurity division. Total revenue reached $487.3 million, representing a 14.2% '
        'year-over-year increase from $426.7 million in FY 2024. This growth was primarily '
        'fueled by a 23% expansion in recurring subscription revenue, which now accounts for '
        '68% of total revenue.'
    ))

    add_body_para(doc, (
        'Operating expenses were carefully managed throughout the year, resulting in an '
        'operating margin improvement from 18.4% to 21.7%. Net income for the period was '
        '$78.2 million, a 28.6% increase over the prior year. Cash flow from operations '
        'totaled $112.5 million, providing ample resources for continued investment in '
        'research and development as well as strategic acquisitions.'
    ))

    add_body_para(doc, (
        'Key accomplishments during the fiscal year include the successful launch of the '
        'Meridian Shield 3.0 cybersecurity platform, the acquisition of DataFlow Analytics '
        'for $45 million, and expansion into the Asia-Pacific market with new offices in '
        'Singapore and Tokyo. Customer retention rate remained exceptional at 94.7%, and '
        'we added 312 net-new enterprise accounts during the period.'
    ))

    add_body_para(doc, (
        'Looking ahead, management remains cautiously optimistic about FY 2026 prospects, '
        'though macroeconomic uncertainties and evolving competitive dynamics warrant careful '
        'monitoring. The strategic initiatives outlined in Section 5 of this report are '
        'designed to position the company for sustained long-term growth.'
    ))

    # === Still PAGE 1 / PAGE 2: Revenue Analysis ===

    add_heading_styled(doc, '1. Revenue Analysis', level=1)

    add_heading_styled(doc, '1.1 Revenue by Segment', level=2)

    add_body_para(doc, (
        'The company operates through three primary business segments: Enterprise Cloud '
        'Solutions, Cybersecurity Products, and Professional Services. Each segment '
        'demonstrated positive growth during the fiscal year, though at varying rates '
        'reflecting different stages of market maturity and competitive positioning.'
    ))

    # Revenue table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['Segment', 'FY 2025 Revenue', 'FY 2024 Revenue', 'YoY Growth']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    data = [
        ['Enterprise Cloud Solutions', '$268.4M', '$224.1M', '19.8%'],
        ['Cybersecurity Products', '$142.7M', '$128.3M', '11.2%'],
        ['Professional Services', '$76.2M', '$74.3M', '2.6%'],
        ['Total', '$487.3M', '$426.7M', '14.2%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if r == 4:
                run.bold = True

    doc.add_paragraph()

    add_body_para(doc, (
        'Enterprise Cloud Solutions continued to be the dominant growth driver, benefiting '
        'from the migration of several large Fortune 500 clients to our platform during Q2 '
        'and Q3. The segment\'s recurring revenue base grew to $198.2 million, representing '
        'a 26.4% increase. Professional Services showed more modest growth, reflecting the '
        'strategic decision to shift from one-time implementation projects to embedded support '
        'within the cloud platform itself.'
    ))

    add_heading_styled(doc, '1.2 Geographic Revenue Distribution', level=2)

    add_body_para(doc, (
        'North America remained the largest geographic market, contributing $358.9 million '
        'or 73.7% of total revenue. However, international markets demonstrated accelerating '
        'growth momentum, with Europe generating $89.4 million (up 18.7%) and Asia-Pacific '
        'contributing $39.0 million (up 42.3%) following the regional expansion.'
    ))

    add_body_para(doc, (
        'The Asia-Pacific growth trajectory was particularly encouraging, as the Singapore '
        'and Tokyo offices secured 47 new enterprise contracts within the first eight months '
        'of operation. Management expects this region to become an increasingly significant '
        'contributor to overall revenue in the coming years, with a target of $65 million '
        'in regional revenue by FY 2027.'
    ))

    # === PAGE 2: Operating Expenses ===

    add_heading_styled(doc, '2. Operating Expenses and Margins', level=1)

    add_heading_styled(doc, '2.1 Cost Structure Overview', level=2)

    add_body_para(doc, (
        'Total operating expenses for FY 2025 were $381.5 million, an increase of 10.8% '
        'from $344.3 million in the prior year. Importantly, expense growth lagged revenue '
        'growth by 3.4 percentage points, reflecting improved operational efficiency and '
        'the benefits of scale in the cloud infrastructure.'
    ))

    # Expense table
    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'
    headers2 = ['Category', 'FY 2025', 'FY 2024', 'Change']
    for i, h in enumerate(headers2):
        cell = table2.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    expense_data = [
        ['Cost of Revenue', '$168.2M', '$152.7M', '+10.2%'],
        ['Research & Development', '$82.4M', '$71.8M', '+14.8%'],
        ['Sales & Marketing', '$78.6M', '$73.5M', '+6.9%'],
        ['General & Administrative', '$52.3M', '$46.3M', '+13.0%'],
        ['Total Operating Expenses', '$381.5M', '$344.3M', '+10.8%'],
    ]
    for r, row_data in enumerate(expense_data, 1):
        for c, val in enumerate(row_data):
            cell = table2.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if r == 5:
                run.bold = True

    doc.add_paragraph()

    add_body_para(doc, (
        'Research and development spending increased by 14.8% as the company invested '
        'heavily in the next-generation Meridian Cloud Platform 4.0 and the artificial '
        'intelligence capabilities being integrated across the product portfolio. Headcount '
        'in the R&D organization grew from 892 to 1,034 employees during the year.'
    ))

    add_body_para(doc, (
        'Sales and marketing expenses grew at a more moderate 6.9%, reflecting improved '
        'sales productivity and the growing contribution of inbound and digital marketing '
        'channels. The sales team closed an average of $2.1 million per quota-carrying '
        'representative, up from $1.8 million in the prior year.'
    ))

    add_heading_styled(doc, '2.2 Margin Analysis', level=2)

    add_body_para(doc, (
        'Gross margin improved to 65.5% from 64.2%, driven primarily by economies of scale '
        'in cloud infrastructure hosting costs and the increasing proportion of high-margin '
        'subscription revenue. Operating margin expanded to 21.7% from 19.3%, while EBITDA '
        'margin reached 28.4%, positioning the company favorably among mid-cap enterprise '
        'software peers.'
    ))

    add_body_para(doc, (
        'Net profit margin improved to 16.0% from 14.3%, reflecting both the operating '
        'leverage described above and a lower effective tax rate of 19.2% compared to '
        '21.5% in the prior year, resulting from the recognition of R&D tax credits and '
        'the favorable tax treatment of intellectual property income in certain jurisdictions.'
    ))

    # === PAGE 3: Forward-Looking Projections ===

    add_heading_styled(doc, '3. Forward-Looking Financial Projections', level=1)

    add_heading_styled(doc, '3.1 Revenue Outlook', level=2)

    add_body_para(doc, (
        'Management has developed financial projections for the upcoming fiscal year based '
        'on current market conditions, the existing sales pipeline, and planned product '
        'launches. These projections incorporate assumptions about customer retention, '
        'new business acquisition, and pricing trends that are subject to change.'
    ))

    # THE TARGET PARAGRAPH - must be on page 3
    add_body_para(doc, (
        'The projected revenue for Q3 is expected to reach $142.8 million, driven by '
        'seasonal enterprise purchasing patterns and the anticipated general availability '
        'of the Meridian Cloud Platform 4.0. This figure assumes a 15% conversion rate '
        'on the current qualified pipeline of $952 million and continued growth in the '
        'existing customer base through upsell and cross-sell motions.'
    ))

    add_body_para(doc, (
        'For the full fiscal year 2026, management projects total revenue in the range of '
        '$548 million to $572 million, representing year-over-year growth of 12.5% to 17.4%. '
        'The breadth of this range reflects uncertainty around the timing of several large '
        'enterprise deals currently in advanced negotiation stages, as well as the pace of '
        'adoption for the new platform release.'
    ))

    add_heading_styled(doc, '3.2 Investment Priorities', level=2)

    add_body_para(doc, (
        'Capital expenditure for FY 2026 is planned at $45 million to $52 million, '
        'primarily directed toward expanding cloud infrastructure capacity in the European '
        'and Asia-Pacific regions. An additional $15 million has been earmarked for potential '
        'strategic acquisitions that would strengthen the company\'s data analytics and '
        'machine learning capabilities.'
    ))

    add_body_para(doc, (
        'The board of directors has approved a share repurchase program of up to $30 million '
        'over the next twelve months, to be executed opportunistically based on market '
        'conditions. The company also intends to maintain its quarterly dividend of $0.15 '
        'per share, reflecting confidence in the sustainability of cash flow generation.'
    ))

    add_heading_styled(doc, '4. Risk Factors', level=1)

    add_body_para(doc, (
        'Several risk factors could materially affect the company\'s financial performance '
        'in the upcoming period. Increased competition from both established enterprise '
        'software vendors and well-funded startups continues to pressure pricing in certain '
        'market segments. Additionally, potential changes to data privacy regulations in '
        'the European Union and certain U.S. states could necessitate significant product '
        'modifications and compliance investments.'
    ))

    add_body_para(doc, (
        'Currency fluctuation risk has increased with the expansion of international '
        'operations. Approximately 26.3% of revenue is now denominated in foreign currencies, '
        'up from 19.1% in the prior year. The company utilizes hedging instruments to mitigate '
        'short-term exposure, but sustained currency movements could impact reported results.'
    ))

    add_body_para(doc, (
        'Talent acquisition and retention remain ongoing challenges, particularly in the '
        'areas of cloud engineering and cybersecurity expertise. Total employee turnover '
        'was 11.2% during the year, slightly below the industry average of 13.5%, but '
        'voluntary attrition among senior technical staff increased to 8.4% from 6.1% in '
        'the prior year, warranting continued attention to compensation and career development '
        'programs.'
    ))

    add_heading_styled(doc, '5. Strategic Initiatives', level=1)

    add_body_para(doc, (
        'Management has identified five strategic priorities for FY 2026: (1) successful '
        'general availability launch of Meridian Cloud Platform 4.0, (2) expansion of the '
        'AI-powered security analytics capability, (3) deepening Asia-Pacific market '
        'penetration, (4) development of an industry-vertical solutions practice for '
        'healthcare and financial services, and (5) integration of the DataFlow Analytics '
        'acquisition into the core product suite.'
    ))

    add_body_para(doc, (
        'Each initiative has been assigned an executive sponsor and a dedicated cross-functional '
        'team. Progress will be tracked through quarterly business reviews, with formal updates '
        'provided to the Board of Directors at each regular meeting. The total investment '
        'allocated to these strategic initiatives is $28 million, funded from operating cash flow.'
    ))

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
