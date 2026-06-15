"""
Initial Setup: Create master file with 6 job descriptions in default formatting
Task ID: writer_hr_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_043'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# 6 job descriptions with realistic content
JOB_DESCRIPTIONS = [
    {
        "title": "Senior Software Engineer",
        "responsibilities": [
            "Design and implement scalable backend services using Python and Go",
            "Lead code reviews and mentor junior developers on best practices",
            "Collaborate with product managers to define technical requirements",
            "Maintain CI/CD pipelines and ensure 99.9% uptime for critical services",
            "Write comprehensive unit and integration tests for all new features",
        ],
        "requirements": [
            "5+ years of experience in software development",
            "Proficiency in Python, Go, or Java with strong OOP fundamentals",
            "Experience with cloud platforms (AWS, GCP, or Azure)",
            "Familiarity with containerization tools such as Docker and Kubernetes",
            "Bachelor's degree in Computer Science or equivalent experience",
        ],
        "benefits": [
            "Competitive salary range of $140,000 - $180,000",
            "Comprehensive health, dental, and vision insurance",
            "401(k) matching up to 6% of annual salary",
            "Flexible remote work policy with quarterly team retreats",
            "Annual learning and development budget of $3,000",
        ],
    },
    {
        "title": "Marketing Manager",
        "responsibilities": [
            "Develop and execute multi-channel marketing campaigns",
            "Manage a team of 4 content creators and 2 graphic designers",
            "Analyze campaign performance using Google Analytics and HubSpot",
            "Coordinate with sales team to align messaging and lead generation goals",
            "Oversee brand guidelines and ensure consistency across all materials",
        ],
        "requirements": [
            "7+ years of marketing experience with at least 3 years in management",
            "Proven track record of managing budgets exceeding $500,000 annually",
            "Strong proficiency in marketing automation platforms",
            "Excellent written and verbal communication skills",
            "MBA or Master's degree in Marketing preferred",
        ],
        "benefits": [
            "Base salary of $110,000 - $135,000 plus performance bonus",
            "Full medical, dental, and vision coverage for employee and dependents",
            "20 days paid time off plus 10 company holidays",
            "Monthly wellness stipend of $150",
            "Stock options vesting over 4 years",
        ],
    },
    {
        "title": "Data Analyst",
        "responsibilities": [
            "Build and maintain dashboards in Tableau and Power BI for executive reporting",
            "Perform exploratory data analysis to identify trends and business opportunities",
            "Collaborate with engineering to define data pipeline requirements",
            "Create weekly and monthly reports for stakeholders across 5 departments",
            "Validate data quality and implement automated data cleaning procedures",
        ],
        "requirements": [
            "3+ years of experience in data analysis or business intelligence",
            "Advanced SQL skills with experience querying large datasets",
            "Proficiency in Python or R for statistical analysis",
            "Experience with visualization tools such as Tableau or Power BI",
            "Strong attention to detail and ability to communicate findings to non-technical audiences",
        ],
        "benefits": [
            "Salary range of $85,000 - $105,000",
            "Health insurance with employer covering 80% of premiums",
            "Flexible working hours with core hours from 10 AM to 3 PM",
            "Professional development reimbursement up to $2,500 per year",
            "Commuter benefits and free parking",
        ],
    },
    {
        "title": "Product Designer",
        "responsibilities": [
            "Create wireframes, prototypes, and high-fidelity mockups using Figma",
            "Conduct user research including interviews, surveys, and usability testing",
            "Collaborate with engineers to ensure pixel-perfect implementation",
            "Maintain and evolve the company design system across web and mobile",
            "Present design rationale and gather feedback from cross-functional stakeholders",
        ],
        "requirements": [
            "4+ years of product design experience in a SaaS environment",
            "Expert-level proficiency in Figma and Adobe Creative Suite",
            "Portfolio demonstrating end-to-end design process for complex workflows",
            "Understanding of accessibility standards (WCAG 2.1 AA)",
            "Experience working in agile development environments",
        ],
        "benefits": [
            "Competitive salary of $100,000 - $130,000",
            "Full benefits package including mental health support",
            "Home office setup allowance of $2,000",
            "Conference attendance budget for up to 2 events per year",
            "Sabbatical leave after 5 years of service",
        ],
    },
    {
        "title": "Financial Controller",
        "responsibilities": [
            "Oversee all accounting operations including accounts payable and receivable",
            "Prepare monthly, quarterly, and annual financial statements",
            "Manage the annual budgeting process and variance analysis",
            "Ensure compliance with GAAP and coordinate external audits",
            "Supervise a team of 6 accounting professionals",
        ],
        "requirements": [
            "8+ years of progressive accounting experience",
            "CPA certification required; CMA preferred",
            "Experience with ERP systems such as SAP or Oracle",
            "Strong knowledge of GAAP, IFRS, and tax regulations",
            "Proven leadership experience managing finance teams of 5 or more",
        ],
        "benefits": [
            "Salary range of $130,000 - $160,000 plus annual bonus",
            "Executive health plan with concierge medical services",
            "25 days paid time off plus floating holidays",
            "Employer-funded pension plan with 8% contribution",
            "Relocation assistance up to $15,000",
        ],
    },
    {
        "title": "Human Resources Business Partner",
        "responsibilities": [
            "Partner with department leaders to develop workforce planning strategies",
            "Lead employee engagement initiatives and culture programs",
            "Manage the full-cycle recruitment process for key positions",
            "Conduct compensation benchmarking and recommend salary adjustments",
            "Handle complex employee relations matters and ensure legal compliance",
        ],
        "requirements": [
            "6+ years of HR experience with HRBP or generalist background",
            "SHRM-CP or PHR certification preferred",
            "Experience with HRIS systems such as Workday or BambooHR",
            "Strong understanding of employment law and regulations",
            "Demonstrated ability to influence senior leadership",
        ],
        "benefits": [
            "Base salary of $95,000 - $120,000",
            "Comprehensive benefits including fertility and adoption support",
            "Hybrid work schedule with 3 days in office",
            "Tuition reimbursement up to $5,250 annually",
            "Employee assistance program and free counseling sessions",
        ],
    },
]


def create_initial():
    doc = Document()

    for i, job in enumerate(JOB_DESCRIPTIONS):
        # Job title - plain paragraph, no special formatting
        title_para = doc.add_paragraph(job["title"])

        # Responsibilities section
        doc.add_paragraph("Responsibilities")
        for item in job["responsibilities"]:
            doc.add_paragraph(item, style="List Bullet")

        # Requirements section
        doc.add_paragraph("Requirements")
        for item in job["requirements"]:
            doc.add_paragraph(item, style="List Bullet")

        # Benefits section
        doc.add_paragraph("Benefits")
        for item in job["benefits"]:
            doc.add_paragraph(item, style="List Bullet")

        # Add separator between jobs (except after the last one)
        if i < len(JOB_DESCRIPTIONS) - 1:
            doc.add_paragraph("")  # blank line separator

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
