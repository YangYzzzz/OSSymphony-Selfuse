"""
Reward Script: Apply character styles to job description documents
Task ID: writer_hr_043
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Three custom character styles exist
  Component 2 (0.30): Job titles formatted (18pt, bold, dark blue)
  Component 3 (0.25): Section headers formatted (14pt, bold, black)
  Component 4 (0.25): Requirement/benefit items formatted (11pt, regular, dark gray)
"""

import os
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_043'

# Expected job titles (6 total)
JOB_TITLES = [
    'Senior Software Engineer',
    'Marketing Manager',
    'Data Analyst',
    'Product Designer',
    'Financial Controller',
    'Human Resources Business Partner',
]

# Section header keywords
SECTION_HEADERS = {'Responsibilities', 'Requirements', 'Benefits'}

# Sections whose bullet items should get Requirements Text formatting
STYLED_SECTIONS = {'Requirements', 'Benefits'}


def color_distance(c1, c2):
    """Euclidean distance between two RGB tuples."""
    return sqrt(sum((a - b) ** 2 for a, b in zip(c1, c2)))


def rgb_tuple(rgb_color):
    """Convert RGBColor to (r, g, b) tuple."""
    return (rgb_color[0], rgb_color[1], rgb_color[2])


def check_runs_format(para, expected_bold, expected_pt, expected_rgb, tolerance_pt=0.5, tolerance_color=30):
    """Check if all runs in a paragraph match the expected formatting.
    Returns True if all runs match, False otherwise."""
    if len(para.runs) == 0:
        return False
    for run in para.runs:
        font = run.font
        # Check bold
        if expected_bold and not font.bold:
            return False
        if not expected_bold and font.bold:
            return False
        # Check size
        if font.size is None or abs(font.size.pt - expected_pt) > tolerance_pt:
            return False
        # Check color
        if font.color and font.color.rgb:
            c = rgb_tuple(font.color.rgb)
            if color_distance(c, expected_rgb) > tolerance_color:
                return False
        else:
            return False
    return True


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ================================================================
    # Component 1: Three custom character styles exist (0.20 points)
    # ================================================================
    try:
        required_styles = ['Job Title Format', 'Section Header Format', 'Requirements Text']
        found_styles = []
        for sname in required_styles:
            try:
                style = doc.styles[sname]
                if style.type is not None and str(style.type) == 'CHARACTER (2)':
                    found_styles.append(sname)
                else:
                    print(f"  WARN: Style '{sname}' exists but type is {style.type}, not CHARACTER")
            except KeyError:
                pass

        if len(found_styles) == 3:
            # Also verify the style definitions have correct properties
            jt_style = doc.styles['Job Title Format']
            sh_style = doc.styles['Section Header Format']
            rt_style = doc.styles['Requirements Text']

            style_checks_passed = 0

            # Job Title Format: 18pt, bold, dark blue (00008B)
            if (jt_style.font.size and abs(jt_style.font.size.pt - 18.0) < 0.5
                    and jt_style.font.bold
                    and jt_style.font.color and jt_style.font.color.rgb
                    and color_distance(rgb_tuple(jt_style.font.color.rgb), (0x00, 0x00, 0x8B)) <= 30):
                style_checks_passed += 1
            else:
                print(f"  WARN: Job Title Format definition incorrect")

            # Section Header Format: 14pt, bold, black (000000)
            if (sh_style.font.size and abs(sh_style.font.size.pt - 14.0) < 0.5
                    and sh_style.font.bold
                    and sh_style.font.color and sh_style.font.color.rgb
                    and color_distance(rgb_tuple(sh_style.font.color.rgb), (0, 0, 0)) <= 30):
                style_checks_passed += 1
            else:
                print(f"  WARN: Section Header Format definition incorrect")

            # Requirements Text: 11pt, regular (not bold), dark gray (404040)
            if (rt_style.font.size and abs(rt_style.font.size.pt - 11.0) < 0.5
                    and rt_style.font.color and rt_style.font.color.rgb
                    and color_distance(rgb_tuple(rt_style.font.color.rgb), (0x40, 0x40, 0x40)) <= 30):
                style_checks_passed += 1
            else:
                print(f"  WARN: Requirements Text definition incorrect")

            if style_checks_passed == 3:
                print(f"PASS: Component 1 - All 3 character styles exist with correct definitions (0.20 pts)")
                total_score += 0.20
            elif style_checks_passed >= 0:
                # Styles exist but definitions are incomplete/wrong
                print(f"PARTIAL: Component 1 - Styles exist but {3 - style_checks_passed} definitions incorrect (0.10 pts)")
                total_score += 0.10
        elif len(found_styles) > 0:
            partial = 0.07 * len(found_styles)
            print(f"PARTIAL: Component 1 - Only {len(found_styles)}/3 character styles found: {found_styles} ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 - No custom character styles found")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # ================================================================
    # Component 2: Job titles formatted correctly (0.30 points)
    # 18pt, bold, dark blue (00008B)
    # ================================================================
    try:
        titles_correct = 0
        titles_found = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text in JOB_TITLES:
                titles_found += 1
                if check_runs_format(para, expected_bold=True, expected_pt=18.0, expected_rgb=(0x00, 0x00, 0x8B)):
                    titles_correct += 1
                else:
                    print(f"  FAIL: Job title '{text}' has incorrect formatting")

        if titles_found == 0:
            print(f"FAIL: Component 2 - No job title paragraphs found")
        elif titles_correct == 6:
            print(f"PASS: Component 2 - All 6 job titles correctly formatted (0.30 pts)")
            total_score += 0.30
        elif titles_correct > 0:
            partial = 0.30 * (titles_correct / 6.0)
            print(f"PARTIAL: Component 2 - {titles_correct}/6 job titles correctly formatted ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 - Found {titles_found} titles but none correctly formatted")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # ================================================================
    # Component 3: Section headers formatted correctly (0.25 points)
    # 14pt, bold, black (000000)
    # Headers: Responsibilities, Requirements, Benefits (3 per job x 6 jobs = 18)
    # ================================================================
    try:
        headers_correct = 0
        headers_found = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text in SECTION_HEADERS:
                headers_found += 1
                if check_runs_format(para, expected_bold=True, expected_pt=14.0, expected_rgb=(0, 0, 0)):
                    headers_correct += 1
                else:
                    print(f"  FAIL: Section header '{text}' has incorrect formatting")

        expected_headers = 18  # 3 sections x 6 jobs
        if headers_found == 0:
            print(f"FAIL: Component 3 - No section header paragraphs found")
        elif headers_correct == headers_found and headers_found >= expected_headers:
            print(f"PASS: Component 3 - All {headers_correct} section headers correctly formatted (0.25 pts)")
            total_score += 0.25
        elif headers_correct > 0:
            partial = 0.25 * (headers_correct / max(expected_headers, headers_found))
            print(f"PARTIAL: Component 3 - {headers_correct}/{headers_found} section headers correct ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 - Found {headers_found} headers but none correctly formatted")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # ================================================================
    # Component 4: Requirement/benefit items formatted (0.25 points)
    # Items under 'Requirements' and 'Benefits' sections: 11pt, not bold, dark gray (404040)
    # ================================================================
    try:
        # Identify which bullet items fall under Requirements or Benefits sections
        # by tracking the current section header
        current_section = None
        items_to_check = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text in JOB_TITLES:
                current_section = None  # reset on new job
            elif text in SECTION_HEADERS:
                current_section = text
            elif para.style and para.style.name == 'List Bullet' and current_section in STYLED_SECTIONS:
                items_to_check.append(para)

        items_correct = 0
        for para in items_to_check:
            if check_runs_format(para, expected_bold=False, expected_pt=11.0, expected_rgb=(0x40, 0x40, 0x40)):
                items_correct += 1

        total_items = len(items_to_check)
        if total_items == 0:
            print(f"FAIL: Component 4 - No requirement/benefit items found")
        elif items_correct == total_items:
            print(f"PASS: Component 4 - All {items_correct}/{total_items} requirement/benefit items correctly formatted (0.25 pts)")
            total_score += 0.25
        elif items_correct > 0:
            partial = 0.25 * (items_correct / total_items)
            print(f"PARTIAL: Component 4 - {items_correct}/{total_items} items correct ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 - Found {total_items} items but none correctly formatted")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Main entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
