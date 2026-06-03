"""
Initial Setup: Multi-chapter litigation brief with title page, TOC placeholder, and four chapter subdocuments.
Task ID: writer_legal_050
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_050'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_chapter(filepath, title, heading_text, body_paragraphs):
    """Create a chapter .docx file with Heading 1 and body content."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    h = doc.add_heading(heading_text, level=1)
    for run in h.runs:
        run.font.name = 'Times New Roman'

    for para_text in body_paragraphs:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.first_line_indent = Inches(0.5)

    doc.save(filepath)
    print(f'Chapter created: {filepath}')


def create_main_document():
    """Create the main document with title page and TOC placeholder."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # --- Title Page ---
    # Add spacing before title
    for _ in range(4):
        doc.add_paragraph('')

    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('LITIGATION BRIEF')
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.name = 'Times New Roman'

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_before = Pt(24)
    sub_run = subtitle.add_run('Westbrook Industries, Inc. v. Meridian Financial Group, LLC')
    sub_run.font.size = Pt(16)
    sub_run.font.name = 'Times New Roman'
    sub_run.italic = True

    case_no = doc.add_paragraph()
    case_no.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    case_no.paragraph_format.space_before = Pt(18)
    cn_run = case_no.add_run('Case No. 2025-CV-04382')
    cn_run.font.size = Pt(14)
    cn_run.font.name = 'Times New Roman'

    court = doc.add_paragraph()
    court.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    court.paragraph_format.space_before = Pt(12)
    ct_run = court.add_run('United States District Court\nEastern District of Virginia')
    ct_run.font.size = Pt(12)
    ct_run.font.name = 'Times New Roman'

    # Attorney info at bottom of title page
    for _ in range(6):
        doc.add_paragraph('')

    atty = doc.add_paragraph()
    atty.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    a_run = atty.add_run(
        'Prepared by:\n'
        'Katherine R. Sullivan, Esq.\n'
        'Sullivan, Park & Associates LLP\n'
        '1200 Commerce Tower, Suite 3400\n'
        'Arlington, VA 22201\n'
        'Tel: (703) 555-0142'
    )
    a_run.font.size = Pt(11)
    a_run.font.name = 'Times New Roman'

    # Page break before TOC
    doc.add_page_break()

    # --- Table of Contents Placeholder ---
    toc_heading = doc.add_heading('TABLE OF CONTENTS', level=1)
    toc_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in toc_heading.runs:
        run.font.name = 'Times New Roman'

    toc_note = doc.add_paragraph()
    toc_note.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    toc_note.paragraph_format.space_before = Pt(36)
    tn_run = toc_note.add_run('[Table of Contents will be generated from subdocument headings]')
    tn_run.italic = True
    tn_run.font.size = Pt(11)
    tn_run.font.name = 'Times New Roman'
    tn_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUTPUT)
    print(f'Main document created: {OUTPUT}')


def create_all_chapters():
    """Create the four chapter subdocuments."""

    # Chapter 1: Statement of Facts
    create_chapter(
        f'{WORKDIR}/Chapter1_Facts.docx',
        'Chapter 1',
        'CHAPTER I: STATEMENT OF FACTS',
        [
            'On March 15, 2024, Westbrook Industries, Inc. ("Westbrook") entered into a '
            'Strategic Partnership Agreement ("the Agreement") with Meridian Financial Group, '
            'LLC ("Meridian") for the provision of comprehensive financial advisory services '
            'in connection with Westbrook\'s planned acquisition of Pacific Northwest '
            'Manufacturing Corp.',

            'Pursuant to Section 4.2 of the Agreement, Meridian was obligated to deliver a '
            'preliminary valuation report within sixty (60) days of the effective date. The '
            'report, delivered on May 28, 2024, contained material misrepresentations regarding '
            'Pacific Northwest Manufacturing\'s outstanding liabilities, specifically omitting '
            'a $14.7 million environmental remediation obligation identified in the target '
            'company\'s most recent 10-K filing.',

            'Westbrook, relying on Meridian\'s valuation, proceeded with the acquisition at '
            'a purchase price of $187.3 million, which exceeded the fair market value by '
            'approximately $22.4 million when the omitted liabilities were properly accounted for.',

            'On September 3, 2024, Westbrook\'s internal audit team discovered the discrepancy '
            'during post-acquisition due diligence. Subsequent investigation revealed that '
            'Meridian\'s lead analyst, David Chen, had been aware of the environmental '
            'liability but failed to include it in the valuation model due to alleged pressure '
            'from Meridian\'s managing partner to present favorable projections.',

            'Westbrook provided written notice of breach to Meridian on October 12, 2024, '
            'pursuant to Section 9.1 of the Agreement. Meridian failed to cure the breach '
            'within the thirty-day cure period specified in Section 9.3.',
        ]
    )

    # Chapter 2: Legal Standard
    create_chapter(
        f'{WORKDIR}/Chapter2_LegalStandard.docx',
        'Chapter 2',
        'CHAPTER II: APPLICABLE LEGAL STANDARD',
        [
            'This action arises under Virginia Code Section 8.01-221.1, which provides a '
            'private right of action for damages resulting from negligent misrepresentation '
            'in the context of professional advisory services. The applicable standard requires '
            'the plaintiff to demonstrate: (1) a false representation of material fact; '
            '(2) made with knowledge of its falsity or reckless disregard for the truth; '
            '(3) with the intent to induce reliance; and (4) resulting in justifiable '
            'reliance and proximate damages.',

            'Under the Supreme Court of Virginia\'s holding in Richmond Investment Partners v. '
            'Commonwealth Trust, 298 Va. 445 (2020), a financial advisor owes a heightened '
            'duty of care when providing valuation opinions that will foreseeably be relied '
            'upon in acquisition decisions exceeding $50 million in total consideration.',

            'Furthermore, the Fourth Circuit has consistently held that the omission of '
            'material environmental liabilities in a due diligence report constitutes '
            'actionable misrepresentation under both federal securities law and state common '
            'law principles. See Coastal Development Corp. v. Atlantic Advisory Group, '
            '841 F.3d 292 (4th Cir. 2018).',

            'The standard of care applicable to financial advisory firms in this jurisdiction '
            'is that of a reasonably competent professional possessing the specialized knowledge '
            'and skills ordinarily employed by members of the profession. National Association '
            'of Financial Advisors Standards of Professional Conduct, Section 3.4 (2023 ed.).',
        ]
    )

    # Chapter 3: Argument
    create_chapter(
        f'{WORKDIR}/Chapter3_Argument.docx',
        'Chapter 3',
        'CHAPTER III: ARGUMENT',
        [
            'A. Meridian Breached Its Contractual Duty of Care',

            'The undisputed evidence demonstrates that Meridian failed to satisfy the '
            'professional standard of care required under Section 4.2 of the Agreement. '
            'The omission of a $14.7 million environmental remediation liability from the '
            'valuation report was not a matter of professional judgment but rather a '
            'fundamental failure to review publicly available regulatory filings.',

            'B. Westbrook Justifiably Relied on Meridian\'s Representations',

            'Westbrook retained Meridian precisely because of its purported expertise in '
            'industrial acquisition valuations. The Agreement explicitly stated that Westbrook '
            'would rely on Meridian\'s analysis in making its acquisition decision. Under '
            'these circumstances, Westbrook\'s reliance was both reasonable and foreseeable.',

            'C. Meridian\'s Negligence Proximately Caused Westbrook\'s Damages',

            'But for Meridian\'s material omission, Westbrook would have either negotiated '
            'a reduced purchase price or declined the acquisition entirely. Expert testimony '
            'from Dr. Patricia Yamamoto, a certified business valuation analyst, confirms that '
            'the $22.4 million overpayment is directly attributable to the omitted liabilities.',

            'D. Westbrook Is Entitled to Consequential Damages',

            'In addition to the $22.4 million in direct overpayment, Westbrook has incurred '
            '$3.8 million in environmental remediation costs that were not anticipated at the '
            'time of acquisition, $1.2 million in additional legal and consulting fees related '
            'to the environmental compliance effort, and $650,000 in lost revenue attributable '
            'to production facility downtime during remediation activities.',
        ]
    )

    # Chapter 4: Conclusion
    create_chapter(
        f'{WORKDIR}/Chapter4_Conclusion.docx',
        'Chapter 4',
        'CHAPTER IV: CONCLUSION',
        [
            'For the foregoing reasons, Westbrook Industries, Inc. respectfully requests that '
            'this Court enter judgment in its favor against Meridian Financial Group, LLC in '
            'the amount of $28,050,000, representing the aggregate of direct damages, '
            'consequential damages, and associated costs incurred as a direct and proximate '
            'result of Meridian\'s negligent misrepresentation.',

            'Westbrook further requests an award of pre-judgment interest at the statutory rate '
            'of six percent (6%) per annum from October 12, 2024, the date on which Meridian '
            'was notified of its breach, through the date of entry of final judgment.',

            'Additionally, Westbrook seeks an award of reasonable attorneys\' fees and costs '
            'incurred in prosecuting this action, as authorized under Section 11.4 of the '
            'Agreement, which provides for fee-shifting in favor of the prevailing party in '
            'any dispute arising under the Agreement.',

            'WHEREFORE, Plaintiff Westbrook Industries, Inc. prays that this Honorable Court '
            'grant the relief requested herein, together with such other and further relief as '
            'this Court deems just and proper.',
        ]
    )


def main():
    create_main_document()
    create_all_chapters()

    # Open the main document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
