"""
Reward Script: Master document with subdocument links for litigation brief
Task ID: writer_legal_050
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Document has multiple sections (>= 4 new section breaks)
  Component 2 (0.25): TOC area contains entries for all 4 chapter titles
  Component 3 (0.35): Four subdocument reference headings present in correct order
  Component 4 (0.15): Subdocument reference headings use Heading 1 style
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_050'

# Expected chapter references in order
CHAPTER_FILES = [
    'Chapter1_Facts.docx',
    'Chapter2_LegalStandard.docx',
    'Chapter3_Argument.docx',
    'Chapter4_Conclusion.docx',
]

CHAPTER_TITLES = [
    'CHAPTER I: STATEMENT OF FACTS',
    'CHAPTER II: APPLICABLE LEGAL STANDARD',
    'CHAPTER III: ARGUMENT',
    'CHAPTER IV: CONCLUSION',
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Document has multiple sections (0.25 points)
    # Initial doc has 1 section; golden should have 5 (title + 4 chapter sections)
    # We check for at least 3 sections as task-introduced change
    try:
        num_sections = len(doc.sections)
        if num_sections >= 5:
            print(f"PASS: Component 1 -- Document has {num_sections} sections (>= 5) (0.25 pts)")
            total_score += 0.25
        elif num_sections >= 3:
            print(f"PARTIAL: Component 1 -- Document has {num_sections} sections (>= 3 but < 5) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 -- Document has only {num_sections} section(s), expected >= 5")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: TOC area contains entries for all 4 chapter titles (0.25 points)
    # The golden doc has TOC entries (P18-P21) with the chapter titles
    # Initial doc only has a placeholder "[Table of Contents will be generated...]"
    try:
        all_text = [p.text.strip() for p in doc.paragraphs]
        full_text = '\n'.join(all_text)

        toc_chapters_found = 0
        for title in CHAPTER_TITLES:
            # Check if the chapter title appears in the document text
            # (as a TOC entry, separate from the subdocument heading sections)
            count = sum(1 for t in all_text if title in t)
            if count >= 1:
                toc_chapters_found += 1

        # We need to verify these appear as TOC entries (before the subdoc sections),
        # not just as subdoc headings. Check that at least some appear near the TOC heading.
        toc_heading_idx = None
        for i, p in enumerate(doc.paragraphs):
            if 'TABLE OF CONTENTS' in p.text.upper():
                toc_heading_idx = i
                break

        toc_entries_found = 0
        if toc_heading_idx is not None:
            # Look in the range after TOC heading but before subdocument sections
            # (roughly within 10 paragraphs of the TOC heading)
            search_end = min(toc_heading_idx + 10, len(doc.paragraphs))
            toc_area_text = [doc.paragraphs[j].text.strip() for j in range(toc_heading_idx + 1, search_end)]
            for title in CHAPTER_TITLES:
                if any(title in t for t in toc_area_text):
                    toc_entries_found += 1

        if toc_entries_found == 4:
            print(f"PASS: Component 2 -- All 4 chapter titles found in TOC area (0.25 pts)")
            total_score += 0.25
        elif toc_entries_found >= 2:
            partial = 0.25 * (toc_entries_found / 4)
            print(f"PARTIAL: Component 2 -- {toc_entries_found}/4 chapter titles in TOC area ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- Only {toc_entries_found}/4 chapter titles found in TOC area")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Four subdocument reference headings present in correct order (0.35 points)
    # Golden doc has headings like "[Content from Chapter1_Facts.docx]" as section markers
    # These represent the subdocument links. Check for references to all 4 chapter files.
    try:
        # Find paragraphs that reference chapter files (as subdocument markers)
        subdoc_refs_found = []
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            for ch_file in CHAPTER_FILES:
                if ch_file in text:
                    subdoc_refs_found.append((i, ch_file))
                    break

        # Check all 4 are present
        files_found = [ref[1] for ref in subdoc_refs_found]
        num_found = len(set(files_found))

        # Check correct order
        correct_order = False
        if num_found == 4:
            expected_order = CHAPTER_FILES
            actual_order = [ref[1] for ref in subdoc_refs_found]
            # Remove duplicates while preserving order
            seen = set()
            unique_order = []
            for f in actual_order:
                if f not in seen:
                    seen.add(f)
                    unique_order.append(f)
            correct_order = unique_order == expected_order

        if num_found == 4 and correct_order:
            print(f"PASS: Component 3 -- All 4 subdocument references found in correct order (0.35 pts)")
            total_score += 0.35
        elif num_found == 4:
            print(f"PARTIAL: Component 3 -- All 4 references found but wrong order (0.20 pts)")
            total_score += 0.20
        elif num_found >= 1:
            partial = 0.35 * (num_found / 4) * 0.7
            print(f"PARTIAL: Component 3 -- {num_found}/4 subdocument references found ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 -- No subdocument references found (looked for Chapter*.docx filenames)")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Subdocument reference headings use Heading 1 style (0.15 points)
    # Checks that the chapter reference paragraphs are styled as headings, not plain text
    try:
        heading_style_count = 0
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            for ch_file in CHAPTER_FILES:
                if ch_file in text:
                    style_name = p.style.name if p.style else 'None'
                    if 'Heading' in style_name:
                        heading_style_count += 1
                    else:
                        print(f"  INFO: Subdoc ref for {ch_file} has style '{style_name}' (not Heading)")
                    break

        if heading_style_count == 4:
            print(f"PASS: Component 4 -- All 4 subdocument headings use Heading style (0.15 pts)")
            total_score += 0.15
        elif heading_style_count >= 2:
            partial = 0.15 * (heading_style_count / 4)
            print(f"PARTIAL: Component 4 -- {heading_style_count}/4 use Heading style ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 -- {heading_style_count}/4 subdocument headings use Heading style")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
