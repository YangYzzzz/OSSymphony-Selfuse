"""
Reward Script: Create a formatted invoice table with merged cells
Task ID: writer_rd_013
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Table exists with 8 rows x 4 cols
  Component 2 (0.20): Row 0 merged across 4 cols, contains 'TechVision Inc.' bold ~16pt
  Component 3 (0.15): Row 1 merged across 4 cols, contains company address
  Component 4 (0.20): Row 2 header with Item/Quantity/Unit Price/Total, bold, gray bg (#D9D9D9)
  Component 5 (0.15): Rows 3-6 contain invoice item data (4 items)
  Component 6 (0.10): Row 7 first 3 cells merged, 'Grand Total:' bold right-aligned, last cell has sum
"""

import os

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_013'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def get_cell_grid_span(cell):
    """Return the gridSpan value for a table cell, or 1 if not set."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        gs_el = tcPr.find(qn('w:gridSpan'))
        if gs_el is not None:
            try:
                return int(gs_el.get(qn('w:val')))
            except (TypeError, ValueError):
                pass
    return 1


def get_cell_shading_fill(cell):
    """Return the fill color hex string (e.g. 'D9D9D9') or None."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            return shd.get(qn('w:fill'))
    return None


def cell_has_bold_text(cell):
    """Check if at least one non-empty run in the cell is bold."""
    for para in cell.paragraphs:
        for run in para.runs:
            if run.text.strip() and run.bold:
                return True
    return False


def get_cell_text(cell):
    """Get stripped text from a cell."""
    return cell.text.strip()


def get_para_alignment(cell):
    """Get alignment of the first paragraph with text."""
    for para in cell.paragraphs:
        if para.text.strip():
            return para.paragraph_format.alignment
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document — task requires inserting an invoice table")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # Component 1: Table has correct dimensions — 8 rows x 4 cols (0.20 points)
    try:
        if num_cols == 4 and num_rows == 8:
            print(f"PASS: Component 1 — Table has {num_rows} rows x {num_cols} cols (0.20 pts)")
            total_score += 0.20
        elif num_cols == 4 and num_rows >= 6:
            # Partial credit: right number of cols, close on rows
            print(f"PARTIAL: Component 1 — Table has {num_rows} rows x {num_cols} cols (expected 8x4) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — Table has {num_rows} rows x {num_cols} cols, expected 8x4")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Row 0 merged across 4 cols with 'TechVision Inc.' bold ~16pt (0.20 points)
    try:
        if num_rows >= 1:
            row0_cell = table.cell(0, 0)
            row0_gs = get_cell_grid_span(row0_cell)
            row0_text = get_cell_text(row0_cell)
            row0_bold = cell_has_bold_text(row0_cell)

            # Check font size ~16pt (203200 EMU = 16pt)
            row0_size_matches = [
                15.0 <= run.font.size.pt <= 17.0
                for para in row0_cell.paragraphs
                for run in para.runs
                if run.text.strip() and run.font.size is not None
            ]
            row0_size_ok = any(row0_size_matches)

            merge_ok = row0_gs >= 4
            text_ok = 'techvision' in row0_text.lower()

            sub_score = 0.0
            if merge_ok:
                sub_score += 0.05
            if text_ok:
                sub_score += 0.05
            if row0_bold:
                sub_score += 0.05
            if row0_size_ok:
                sub_score += 0.05

            if sub_score > 0:
                total_score += sub_score
                details = f"merge={merge_ok}, text='{row0_text[:30]}', bold={row0_bold}, size_ok={row0_size_ok}"
                print(f"PASS: Component 2 — Row 0 company header: {details} ({sub_score:.2f} pts)")
            else:
                print(f"FAIL: Component 2 — Row 0 not properly configured")
        else:
            print("FAIL: Component 2 — Table has no rows")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Row 1 merged across 4 cols with company address (0.15 points)
    try:
        if num_rows >= 2:
            row1_cell = table.cell(1, 0)
            row1_gs = get_cell_grid_span(row1_cell)
            row1_text = get_cell_text(row1_cell)

            merge_ok = row1_gs >= 4
            # Address should contain some address-like content
            has_address = len(row1_text) > 10 and any(w in row1_text.lower() for w in ['boulevard', 'street', 'drive', 'ave', 'road', 'tx', 'ca', 'ny', 'innovation'])

            sub_score = 0.0
            if merge_ok:
                sub_score += 0.075
            if has_address:
                sub_score += 0.075

            if sub_score > 0:
                total_score += sub_score
                print(f"PASS: Component 3 — Row 1 address: merge={merge_ok}, has_address={has_address}, text='{row1_text[:50]}' ({sub_score:.3f} pts)")
            else:
                print(f"FAIL: Component 3 — Row 1 not properly configured")
        else:
            print("FAIL: Component 3 — Table has < 2 rows")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Row 2 is header row with Item/Quantity/Unit Price/Total, bold, gray bg (0.20 points)
    try:
        if num_rows >= 3:
            expected_headers = ['item', 'quantity', 'unit price', 'total']
            headers_found = 0
            bold_count = 0
            gray_bg_count = 0

            for ci in range(min(num_cols, 4)):
                cell = table.cell(2, ci)
                text = get_cell_text(cell).lower()
                if ci < len(expected_headers) and expected_headers[ci] in text:
                    headers_found += 1
                if cell_has_bold_text(cell):
                    bold_count += 1
                fill = get_cell_shading_fill(cell)
                if fill and fill.upper() == 'D9D9D9':
                    gray_bg_count += 1

            sub_score = 0.0
            if headers_found >= 4:
                sub_score += 0.08
            elif headers_found >= 2:
                sub_score += 0.04
            if bold_count >= 4:
                sub_score += 0.06
            elif bold_count >= 2:
                sub_score += 0.03
            if gray_bg_count >= 4:
                sub_score += 0.06
            elif gray_bg_count >= 2:
                sub_score += 0.03

            if sub_score > 0:
                total_score += sub_score
                print(f"PASS: Component 4 — Header row: headers={headers_found}/4, bold={bold_count}/4, gray_bg={gray_bg_count}/4 ({sub_score:.2f} pts)")
            else:
                print(f"FAIL: Component 4 — Header row not properly configured")
        else:
            print("FAIL: Component 4 — Table has < 3 rows")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Rows 3-6 contain invoice item data — 4 items (0.15 points)
    try:
        if num_rows >= 7:
            items_with_data = 0
            for ri in range(3, 7):
                # An item row should have text in col 0 (item name), col 1 (qty), col 3 (total)
                item_name = get_cell_text(table.cell(ri, 0))
                qty_text = get_cell_text(table.cell(ri, 1))
                total_text = get_cell_text(table.cell(ri, 3)) if num_cols > 3 else ""

                if item_name and qty_text and total_text:
                    items_with_data += 1

            sub_score = 0.0
            if items_with_data >= 4:
                sub_score = 0.15
            elif items_with_data >= 2:
                sub_score = 0.075

            if sub_score > 0:
                total_score += sub_score
                print(f"PASS: Component 5 — Invoice items: {items_with_data}/4 rows with data ({sub_score:.3f} pts)")
            else:
                print(f"FAIL: Component 5 — Only {items_with_data}/4 item rows have data")
        else:
            print(f"FAIL: Component 5 — Table has {num_rows} rows, need >= 7 for item rows")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Row 7 — first 3 cells merged, 'Grand Total:' bold right-aligned, last cell has sum (0.10 points)
    try:
        if num_rows >= 8:
            row7_cell0 = table.cell(7, 0)
            row7_gs = get_cell_grid_span(row7_cell0)
            row7_text = get_cell_text(row7_cell0)
            row7_bold = cell_has_bold_text(row7_cell0)

            # Check right alignment
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            row7_align = get_para_alignment(row7_cell0)
            row7_right = (row7_align == WD_PARAGRAPH_ALIGNMENT.RIGHT)

            # Check last cell has a monetary value
            row7_last = get_cell_text(table.cell(7, 3))
            has_total_value = '$' in row7_last or any(c.isdigit() for c in row7_last)

            merge_ok = row7_gs >= 3
            text_ok = 'grand total' in row7_text.lower()

            sub_score = 0.0
            if merge_ok and text_ok:
                sub_score += 0.04
            if row7_bold:
                sub_score += 0.02
            if row7_right:
                sub_score += 0.02
            if has_total_value:
                sub_score += 0.02

            if sub_score > 0:
                total_score += sub_score
                print(f"PASS: Component 6 — Grand Total row: merge={merge_ok}, text_ok={text_ok}, bold={row7_bold}, right={row7_right}, value='{row7_last}' ({sub_score:.2f} pts)")
            else:
                print(f"FAIL: Component 6 — Grand Total row not properly configured")
        else:
            print(f"FAIL: Component 6 — Table has {num_rows} rows, need 8 for Grand Total row")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
