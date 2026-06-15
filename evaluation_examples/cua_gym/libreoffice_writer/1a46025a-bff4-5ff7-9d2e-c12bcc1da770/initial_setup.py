"""
Initial Setup: Create a Writer document with title 'Invoice #2025-0042' and no tables.
Task ID: writer_rd_013
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_013'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set A4 page size with 2.54 cm margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Title
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(18)
    title_run = title_para.add_run('Invoice #2025-0042')
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.name = 'Arial'

    # Date line
    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date_run = date_para.add_run('Date: March 15, 2025')
    date_run.font.size = Pt(11)
    date_run.font.name = 'Arial'

    # Bill To section
    bill_para = doc.add_paragraph()
    bill_para.paragraph_format.space_before = Pt(12)
    bill_run = bill_para.add_run('Bill To:')
    bill_run.bold = True
    bill_run.font.size = Pt(11)
    bill_run.font.name = 'Arial'

    client_para = doc.add_paragraph()
    client_run = client_para.add_run('Meridian Solutions Ltd.\n45 Commerce Drive, Suite 200\nSan Francisco, CA 94105')
    client_run.font.size = Pt(11)
    client_run.font.name = 'Arial'

    # Instruction placeholder
    note_para = doc.add_paragraph()
    note_para.paragraph_format.space_before = Pt(24)
    note_run = note_para.add_run('Please insert the invoice table below this line.')
    note_run.font.size = Pt(11)
    note_run.font.name = 'Arial'
    note_run.italic = True

    # No tables - this is the initial state before the agent acts

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
