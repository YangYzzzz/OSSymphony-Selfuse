"""
Reward Script: Training manual chapter in LibreOffice Writer
Task ID: writer_wf_015
Domain: libreoffice_writer
Scoring:
  Component 1 — Title heading "Chapter 5: Customer Service Best Practices" (0.15)
  Component 2 — TOC placeholder present (0.10)
  Component 3 — Four Heading 2 sections with correct names (0.25)
  Component 4 — Six numbered steps in Handling Complaints (0.20)
  Component 5 — TIP text box (table) with green border (0.15)
  Component 6 — 6pt paragraph after-spacing (0.15)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_015'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gate: document must have content (initial blank doc has 0 paragraphs)
    if len(doc.paragraphs) == 0:
        print("FAIL: Document has no paragraphs — blank document")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title heading — Heading 1 with "Chapter 5: Customer Service Best Practices" (0.15 pts)
    try:
        h1_paras = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
        title_found = any(
            "chapter 5" in p.text.lower() and "customer service" in p.text.lower()
            for p in h1_paras
        )
        if title_found:
            print(f"PASS: Component 1 — Title heading found ({h1_paras[0].text!r}) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — No Heading 1 with 'Chapter 5' and 'Customer Service'. H1s found: {[p.text for p in h1_paras]}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: TOC placeholder present (0.10 pts)
    # The task says "Include: a TOC". In golden, there's a paragraph referencing table of contents.
    try:
        toc_found = any(
            "table of contents" in p.text.lower() or "toc" in p.text.lower()
            for p in doc.paragraphs
        )
        # Also check for TOC field codes in XML (real TOC uses w:fldChar with TOC instruction)
        if not toc_found:
            body_xml = doc.element.body.xml
            toc_found = "TOC" in body_xml and "fldChar" in body_xml
        if toc_found:
            print(f"PASS: Component 2 — TOC reference found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — No TOC or TOC placeholder found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Four Heading 2 sections with correct names (0.25 pts)
    # Expected: Communication Skills, Handling Complaints, Escalation Procedures, Documentation Requirements
    try:
        h2_texts = [p.text.strip() for p in doc.paragraphs if p.style.name == "Heading 2"]
        expected_h2 = [
            "Communication Skills",
            "Handling Complaints",
            "Escalation Procedures",
            "Documentation Requirements",
        ]
        matched = 0
        for exp in expected_h2:
            if any(exp.lower() in h.lower() for h in h2_texts):
                matched += 1
        # Award partial: 0.0625 per matched heading
        component_score = (matched / 4.0) * 0.25
        if matched == 4:
            print(f"PASS: Component 3 — All 4 Heading 2 sections found: {h2_texts} (0.25 pts)")
        else:
            print(f"PARTIAL: Component 3 — {matched}/4 Heading 2 sections found: {h2_texts} ({component_score:.3f} pts)")
        total_score += component_score
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Six numbered steps in Handling Complaints section (0.20 pts)
    # Steps should use List Number style and appear after the "Handling Complaints" heading
    try:
        # Find paragraphs between "Handling Complaints" heading and next heading
        in_complaints = False
        numbered_steps = 0
        for p in doc.paragraphs:
            if p.style.name == "Heading 2" and "handling complaints" in p.text.lower():
                in_complaints = True
                continue
            if in_complaints and p.style.name.startswith("Heading"):
                break
            if in_complaints and "List Number" in p.style.name:
                numbered_steps += 1

        if numbered_steps >= 6:
            print(f"PASS: Component 4 — {numbered_steps} numbered steps in Handling Complaints (0.20 pts)")
            total_score += 0.20
        elif numbered_steps > 0:
            component_score = (min(numbered_steps, 6) / 6.0) * 0.20
            print(f"PARTIAL: Component 4 — {numbered_steps}/6 numbered steps found ({component_score:.3f} pts)")
            total_score += component_score
        else:
            print(f"FAIL: Component 4 — No numbered steps in Handling Complaints section")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: TIP text box with green border (0.15 pts)
    # Implemented as a table with green borders containing "TIP" text
    try:
        tip_table_found = False
        green_border_found = False

        for table in doc.tables:
            # Check if any cell contains "TIP"
            has_tip = False
            for row in table.rows:
                for cell in row.cells:
                    if "tip" in cell.text.lower():
                        has_tip = True
                        break
                if has_tip:
                    break

            if has_tip:
                tip_table_found = True
                # Check for green border on the table
                tbl_elem = table._tbl
                tbl_pr = tbl_elem.find(qn("w:tblPr"))
                if tbl_pr is not None:
                    borders = tbl_pr.find(qn("w:tblBorders"))
                    if borders is not None:
                        for child in borders:
                            color = child.get(qn("w:color"))
                            if color:
                                color_lower = color.lower()
                                # Accept various shades of green
                                # 00B050 is the exact color used, but also accept other greens
                                r = int(color_lower[0:2], 16) if len(color_lower) >= 6 else 0
                                g = int(color_lower[2:4], 16) if len(color_lower) >= 6 else 0
                                b = int(color_lower[4:6], 16) if len(color_lower) >= 6 else 0
                                if g > r and g > b and g >= 128:
                                    green_border_found = True
                                    break

        if tip_table_found and green_border_found:
            print(f"PASS: Component 5 — TIP box with green border found (0.15 pts)")
            total_score += 0.15
        elif tip_table_found:
            print(f"PARTIAL: Component 5 — TIP box found but no green border (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 — No TIP text box found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: 6pt paragraph after-spacing (0.15 pts)
    # Task says "Set paragraph spacing to 6pt after each paragraph"
    # Check that the majority of paragraphs have space_after == 76200 EMU (6pt)
    try:
        total_paras = len(doc.paragraphs)
        paras_with_6pt = 0
        for p in doc.paragraphs:
            sa = p.paragraph_format.space_after
            if sa is not None and abs(sa - 76200) < 1000:  # ~6pt with small tolerance
                paras_with_6pt += 1

        if total_paras > 0:
            ratio = paras_with_6pt / total_paras
            if ratio >= 0.8:
                print(f"PASS: Component 6 — {paras_with_6pt}/{total_paras} paragraphs have 6pt after-spacing (0.15 pts)")
                total_score += 0.15
            elif ratio >= 0.5:
                component_score = 0.10
                print(f"PARTIAL: Component 6 — {paras_with_6pt}/{total_paras} paragraphs have 6pt spacing ({component_score} pts)")
                total_score += component_score
            else:
                print(f"FAIL: Component 6 — Only {paras_with_6pt}/{total_paras} paragraphs have 6pt after-spacing")
        else:
            print(f"FAIL: Component 6 — No paragraphs to check")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: persist app state then verify
def persist_app_state(domain):
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
