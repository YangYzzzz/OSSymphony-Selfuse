"""
Initial Setup: Retainer agreement template with bracketed placeholders
Task ID: writer_legal_074
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_074'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title ---
    title = doc.add_heading('LEGAL RETAINER AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Preamble ---
    preamble = doc.add_paragraph()
    preamble.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    preamble.add_run(
        'This Retainer Agreement ("Agreement") is entered into as of '
    )
    run_date = preamble.add_run('[EFFECTIVE_DATE]')
    run_date.bold = True
    preamble.add_run(
        ' by and between Hartwell & Associates LLP, a law firm organized '
        'and existing under the laws of the State of California ("Firm"), '
        'and '
    )
    run_client = preamble.add_run('[CLIENT_NAME]')
    run_client.bold = True
    preamble.add_run(' ("Client"), located at ')
    run_addr = preamble.add_run('[CLIENT_ADDRESS]')
    run_addr.bold = True
    preamble.add_run('.')

    doc.add_paragraph()  # blank line

    # --- Section 1: Scope of Services ---
    h1 = doc.add_heading('1. SCOPE OF SERVICES', level=1)
    p1 = doc.add_paragraph()
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p1.add_run(
        'The Firm agrees to provide legal services to the Client in connection '
        'with general corporate matters, contract review, regulatory compliance, '
        'and such other legal matters as may be mutually agreed upon from time to '
        'time. The scope of services may be modified by written agreement of both parties.'
    )

    # --- Section 2: Retainer Fee ---
    h2 = doc.add_heading('2. RETAINER FEE', level=1)
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p2.add_run(
        'The Client agrees to pay a retainer fee of '
    )
    run_amount = p2.add_run('[RETAINER_AMOUNT]')
    run_amount.bold = True
    p2.add_run(
        ' (the "Retainer") upon execution of this Agreement. The Retainer '
        'shall be deposited into the Firm\'s client trust account and applied '
        'against fees and costs incurred in the representation. The Firm shall '
        'provide monthly statements detailing services rendered and amounts '
        'deducted from the Retainer.'
    )

    p2a = doc.add_paragraph()
    p2a.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p2a.add_run(
        'Should the Retainer balance fall below twenty-five percent (25%) of '
        'the initial deposit, the Client agrees to replenish the Retainer to '
        'its original amount within fifteen (15) business days of receiving '
        'written notice from the Firm.'
    )

    # --- Section 3: Billing Rate ---
    h3 = doc.add_heading('3. BILLING RATE AND EXPENSES', level=1)
    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p3.add_run(
        'Legal services shall be billed at the rate of '
    )
    run_rate = p3.add_run('[BILLING_RATE]')
    run_rate.bold = True
    p3.add_run(
        ' per hour. This rate is subject to annual review and may be adjusted '
        'with sixty (60) days\' prior written notice to the Client. In addition '
        'to attorney fees, the Client shall reimburse the Firm for all '
        'reasonable out-of-pocket expenses incurred in the course of '
        'representation, including but not limited to court filing fees, '
        'expert witness fees, travel expenses, copying charges, and postage.'
    )

    # --- Section 4: Term and Termination ---
    h4 = doc.add_heading('4. TERM AND TERMINATION', level=1)
    p4 = doc.add_paragraph()
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p4.add_run(
        'This Agreement shall commence on the Effective Date and continue in '
        'effect until terminated by either party. Either party may terminate '
        'this Agreement at any time by providing thirty (30) days\' written '
        'notice to the other party. Upon termination, the Firm shall render a '
        'final statement of account, and any unused portion of the Retainer '
        'shall be refunded to the Client within thirty (30) days.'
    )

    # --- Section 5: Confidentiality ---
    h5 = doc.add_heading('5. CONFIDENTIALITY', level=1)
    p5 = doc.add_paragraph()
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p5.add_run(
        'All communications between the Client and the Firm shall be deemed '
        'privileged and confidential under the attorney-client privilege and '
        'applicable rules of professional conduct. The Firm shall not disclose '
        'any confidential information without the Client\'s prior written '
        'consent, except as required by law or court order.'
    )

    # --- Section 6: Governing Law ---
    h6 = doc.add_heading('6. GOVERNING LAW', level=1)
    p6 = doc.add_paragraph()
    p6.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p6.add_run(
        'This Agreement shall be governed by and construed in accordance with '
        'the laws of the State of California, without regard to its conflict '
        'of laws provisions. Any dispute arising under this Agreement shall be '
        'resolved through binding arbitration in San Francisco County, California.'
    )

    # --- Signature Block ---
    doc.add_paragraph()  # spacing
    sig_header = doc.add_heading('SIGNATURES', level=2)

    p_firm = doc.add_paragraph()
    p_firm.add_run('FIRM: ').bold = True
    p_firm.add_run('Hartwell & Associates LLP')

    p_by = doc.add_paragraph()
    p_by.add_run('By: ________________________________')

    p_name_firm = doc.add_paragraph()
    p_name_firm.add_run('Name: Victoria Hartwell, Managing Partner')

    p_date_firm = doc.add_paragraph()
    p_date_firm.add_run('Date: ________________________________')

    doc.add_paragraph()  # spacing

    p_client_sig = doc.add_paragraph()
    p_client_sig.add_run('CLIENT: ').bold = True
    run_cn = p_client_sig.add_run('[CLIENT_NAME]')
    run_cn.bold = True

    p_by_client = doc.add_paragraph()
    p_by_client.add_run('By: ________________________________')

    p_date_client = doc.add_paragraph()
    p_date_client.add_run('Date: ________________________________')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
