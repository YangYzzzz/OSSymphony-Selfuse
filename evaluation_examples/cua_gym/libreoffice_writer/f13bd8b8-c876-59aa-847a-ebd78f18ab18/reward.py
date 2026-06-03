"""
Reward Script: Research Proposal in LibreOffice Writer
Task ID: writer_wf_066
Domain: libreoffice_writer
Scoring:
  C1: Cover page (title + PI name + centered)  — 0.15
  C2: Table of Contents section present         — 0.10
  C3: 8 Heading 1 sections                      — 0.20
  C4: 3 Methodology subsections (Heading 2)     — 0.10
  C5: 3 numbered research questions             — 0.10
  C6: Budget table (header + 5 items, 3 cols)   — 0.15
  C7: Timeline table (4 phases, Gantt-style)    — 0.10
  C8: 4 references                              — 0.10
  Total: 1.0
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_066'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather all paragraph data once
    all_paras = doc.paragraphs
    heading1_texts = [p.text.strip() for p in all_paras if p.style and p.style.name == 'Heading 1']
    heading2_texts = [p.text.strip() for p in all_paras if p.style and p.style.name == 'Heading 2']

    # ============================================================
    # Component 1: Cover page — title and PI name present & centered (0.15 pts)
    # ============================================================
    try:
        title_found = False
        pi_found = False
        for p in all_paras:
            txt = p.text.strip().lower()
            alignment = p.paragraph_format.alignment
            is_centered = (alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
            if 'urban heat island' in txt and 'coastal' in txt and is_centered:
                title_found = True
            if 'dr. lisa chen' in txt and is_centered:
                pi_found = True

        if title_found and pi_found:
            print(f"PASS: Component 1 — Cover page has centered title and PI name (0.15 pts)")
            total_score += 0.15
        elif title_found or pi_found:
            print(f"PARTIAL: Component 1 — Title found={title_found}, PI found={pi_found} (0.07 pts)")
            total_score += 0.07
        else:
            print(f"FAIL: Component 1 — No centered title or PI name found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ============================================================
    # Component 2: Table of Contents section present (0.10 pts)
    # ============================================================
    try:
        toc_heading_found = any('table of contents' in h.lower() for h in heading1_texts)
        # Also check for TOC-like entries (numbered section references)
        toc_entries = 0
        for p in all_paras:
            txt = p.text.strip()
            # TOC entries typically have tab-separated section names with page numbers
            if '\t' in txt and any(kw in txt.lower() for kw in ['introduction', 'methodology', 'budget', 'references']):
                toc_entries += 1

        if toc_heading_found and toc_entries >= 3:
            print(f"PASS: Component 2 — TOC heading + {toc_entries} TOC entries (0.10 pts)")
            total_score += 0.10
        elif toc_heading_found:
            print(f"PARTIAL: Component 2 — TOC heading found but only {toc_entries} entries (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 2 — No TOC section found in Heading 1 styles")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ============================================================
    # Component 3: 8 Heading 1 sections (0.20 pts)
    # Expected: Table of Contents, Introduction, Literature Review,
    #           Research Questions, Methodology, Expected Outcomes,
    #           Budget, Timeline, References
    # ============================================================
    try:
        expected_h1 = [
            'introduction', 'literature review', 'research questions',
            'methodology', 'expected outcomes', 'budget', 'timeline', 'references'
        ]
        found_h1 = [e for e in expected_h1 if any(e in h.lower() for h in heading1_texts)]
        ratio = len(found_h1) / len(expected_h1)

        if ratio >= 1.0:
            print(f"PASS: Component 3 — All 8 Heading 1 sections found: {found_h1} (0.20 pts)")
            total_score += 0.20
        elif ratio >= 0.5:
            pts = round(0.20 * ratio, 2)
            print(f"PARTIAL: Component 3 — {len(found_h1)}/8 sections found: {found_h1} ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 3 — Only {len(found_h1)}/8 sections: {found_h1}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ============================================================
    # Component 4: Methodology subsections — 3 Heading 2s (0.10 pts)
    # Expected: Study Sites, Data Collection, Analysis
    # ============================================================
    try:
        expected_h2 = ['study sites', 'data collection', 'analysis']
        found_h2 = [e for e in expected_h2 if any(e in h.lower() for h in heading2_texts)]

        if len(found_h2) >= 3:
            print(f"PASS: Component 4 — All 3 methodology subsections found: {found_h2} (0.10 pts)")
            total_score += 0.10
        elif len(found_h2) >= 1:
            pts = round(0.10 * len(found_h2) / 3, 2)
            print(f"PARTIAL: Component 4 — {len(found_h2)}/3 subsections: {found_h2} ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 4 — No methodology subsections found in Heading 2")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ============================================================
    # Component 5: 3 numbered research questions (0.10 pts)
    # ============================================================
    try:
        list_number_paras = [p for p in all_paras if p.style and p.style.name == 'List Number' and p.text.strip()]
        # Also detect manually numbered questions (1. 2. 3.)
        manual_numbered = []
        in_rq_section = False
        for p in all_paras:
            if p.style and p.style.name == 'Heading 1' and 'research question' in p.text.lower():
                in_rq_section = True
                continue
            if p.style and p.style.name == 'Heading 1' and in_rq_section:
                break
            if in_rq_section and p.text.strip():
                # Check if paragraph is a question (contains '?')
                if '?' in p.text:
                    manual_numbered.append(p.text.strip()[:60])

        num_questions = max(len(list_number_paras), len(manual_numbered))

        if num_questions >= 3:
            print(f"PASS: Component 5 — {num_questions} numbered research questions found (0.10 pts)")
            total_score += 0.10
        elif num_questions >= 1:
            pts = round(0.10 * num_questions / 3, 2)
            print(f"PARTIAL: Component 5 — {num_questions}/3 questions ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 5 — No numbered research questions found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # ============================================================
    # Component 6: Budget table — header + 5 data rows, 3 cols (0.15 pts)
    # ============================================================
    try:
        budget_table = None
        for t in doc.tables:
            # Look for a table with budget-related headers
            if len(t.rows) >= 2:
                header_text = ' '.join(cell.text.strip().lower() for cell in t.rows[0].cells)
                if 'item' in header_text and ('cost' in header_text or 'amount' in header_text):
                    budget_table = t
                    break

        if budget_table:
            num_data_rows = len(budget_table.rows) - 1  # exclude header
            num_cols = len(budget_table.columns)
            has_cost_data = False
            for row in list(budget_table.rows)[1:]:
                cell_text = row.cells[1].text.strip() if len(row.cells) > 1 else ''
                if '$' in cell_text or cell_text.replace(',', '').replace('.', '').isdigit():
                    has_cost_data = True
                    break

            if num_data_rows >= 5 and num_cols >= 3 and has_cost_data:
                print(f"PASS: Component 6 — Budget table: {num_data_rows} data rows, {num_cols} cols, has cost data (0.15 pts)")
                total_score += 0.15
            elif num_data_rows >= 3 and num_cols >= 3:
                print(f"PARTIAL: Component 6 — Budget table: {num_data_rows} rows, {num_cols} cols (0.08 pts)")
                total_score += 0.08
            else:
                print(f"FAIL: Component 6 — Budget table has {num_data_rows} data rows, {num_cols} cols")
        else:
            print(f"FAIL: Component 6 — No budget table found with Item/Cost headers")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # ============================================================
    # Component 7: Timeline/Gantt table — 4 phases (0.10 pts)
    # ============================================================
    try:
        timeline_table = None
        for t in doc.tables:
            if len(t.rows) >= 2:
                header_text = ' '.join(cell.text.strip().lower() for cell in t.rows[0].cells)
                if 'phase' in header_text and ('month' in header_text or 'quarter' in header_text or 'q1' in header_text):
                    timeline_table = t
                    break

        if timeline_table:
            num_phases = 0
            for row in list(timeline_table.rows)[1:]:
                cell_text = row.cells[0].text.strip().lower()
                if 'phase' in cell_text:
                    num_phases += 1

            num_cols = len(timeline_table.columns)
            # Check for Gantt markers (X or filled cells)
            has_markers = False
            for row in list(timeline_table.rows)[1:]:
                for cell in list(row.cells)[1:]:
                    if cell.text.strip():
                        has_markers = True
                        break
                if has_markers:
                    break

            if num_phases >= 4 and has_markers:
                print(f"PASS: Component 7 — Timeline table: {num_phases} phases, {num_cols} cols, has markers (0.10 pts)")
                total_score += 0.10
            elif num_phases >= 2:
                pts = round(0.10 * num_phases / 4, 2)
                print(f"PARTIAL: Component 7 — {num_phases}/4 phases ({pts} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 7 — Timeline table has only {num_phases} phases")
        else:
            print(f"FAIL: Component 7 — No timeline table found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # ============================================================
    # Component 8: References section with 4 entries (0.10 pts)
    # ============================================================
    try:
        in_references = False
        ref_entries = []
        for p in all_paras:
            if p.style and p.style.name == 'Heading 1' and 'reference' in p.text.lower():
                in_references = True
                continue
            if p.style and p.style.name == 'Heading 1' and in_references:
                break
            if in_references and p.text.strip():
                ref_entries.append(p.text.strip()[:60])

        if len(ref_entries) >= 4:
            print(f"PASS: Component 8 — {len(ref_entries)} references found (0.10 pts)")
            total_score += 0.10
        elif len(ref_entries) >= 2:
            pts = round(0.10 * len(ref_entries) / 4, 2)
            print(f"PARTIAL: Component 8 — {len(ref_entries)}/4 references ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 8 — Only {len(ref_entries)} references found")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
