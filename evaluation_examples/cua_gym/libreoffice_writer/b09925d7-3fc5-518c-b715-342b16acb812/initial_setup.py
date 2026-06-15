"""
Initial Setup: Create a 10-section contract template with no protected sections.
Task ID: writer_legal_048
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_048'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Contract No. PSA-2025-0472')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph('')  # spacer

    # --- Section 1: Preamble ---
    doc.add_heading('Section 1: Preamble', level=1)
    doc.add_paragraph(
        'This Professional Services Agreement ("Agreement") is entered into as of '
        'March 15, 2025 ("Effective Date"), by and between Meridian Technology Solutions, Inc., '
        'a Delaware corporation with principal offices at 2400 Innovation Drive, Suite 800, '
        'San Jose, California 95134 ("Service Provider"), and Cascade Financial Group, LLC, '
        'a New York limited liability company with principal offices at 180 Park Avenue, '
        '22nd Floor, New York, New York 10166 ("Client").'
    )
    doc.add_paragraph(
        'WHEREAS, the Client desires to engage the Service Provider to perform certain '
        'professional consulting and technology integration services; and WHEREAS, the '
        'Service Provider has the expertise, personnel, and resources to provide such services; '
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth herein, '
        'and for other good and valuable consideration, the receipt and sufficiency of which are '
        'hereby acknowledged, the parties agree as follows.'
    )

    # --- Section 2: Definitions ---
    doc.add_heading('Section 2: Definitions', level=1)
    doc.add_paragraph(
        '"Confidential Information" means any non-public information disclosed by either party '
        'to the other, whether orally, in writing, or by inspection, including but not limited to '
        'business plans, financial data, customer lists, technical specifications, trade secrets, '
        'and proprietary methodologies.'
    )
    doc.add_paragraph(
        '"Deliverables" means all work products, reports, software, documentation, and other '
        'materials developed or produced by the Service Provider in the course of performing '
        'the Services, as further described in Exhibit A attached hereto.'
    )
    doc.add_paragraph(
        '"Services" means the professional consulting, technology integration, and advisory '
        'services to be performed by the Service Provider as described in Section 3 and '
        'detailed in the Statement of Work ("SOW") attached as Exhibit B.'
    )

    # --- Section 3: Scope of Services ---
    doc.add_heading('Section 3: Scope of Services', level=1)
    doc.add_paragraph(
        'The Service Provider shall perform the following services for the Client during '
        'the term of this Agreement:'
    )
    doc.add_paragraph(
        'a) Conduct a comprehensive assessment of the Client\'s existing enterprise resource '
        'planning (ERP) infrastructure and provide detailed recommendations for modernization.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'b) Design and implement a cloud-based data integration platform connecting the '
        'Client\'s CRM, financial, and operations systems.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'c) Provide training sessions for up to 50 Client personnel on the new systems, '
        'including user manuals and reference guides.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'd) Deliver monthly progress reports and participate in quarterly executive reviews '
        'throughout the engagement period.',
        style='List Bullet'
    )

    # --- Section 4: Compensation and Payment ---
    doc.add_heading('Section 4: Compensation and Payment', level=1)
    doc.add_paragraph(
        'In consideration of the Services, the Client shall pay the Service Provider a total '
        'fixed fee of $475,000.00, payable in the following installments:'
    )
    doc.add_paragraph(
        'Phase 1 - Assessment & Planning: $95,000.00, due upon execution of this Agreement.'
    )
    doc.add_paragraph(
        'Phase 2 - Design & Development: $190,000.00, payable in two equal installments of '
        '$95,000.00 each, due at the completion of design review and midpoint development milestones.'
    )
    doc.add_paragraph(
        'Phase 3 - Implementation & Training: $142,500.00, due upon successful completion '
        'of user acceptance testing.'
    )
    doc.add_paragraph(
        'Phase 4 - Post-Implementation Support: $47,500.00, due at the conclusion of the '
        '90-day support period.'
    )
    doc.add_paragraph(
        'All invoices shall be payable within thirty (30) calendar days of receipt. '
        'Late payments shall accrue interest at the rate of 1.5% per month or the maximum '
        'rate permitted by law, whichever is less.'
    )

    # --- Section 5: Standard Terms and Conditions ---
    doc.add_heading('Section 5: Standard Terms and Conditions', level=1)
    doc.add_paragraph(
        'The Service Provider shall perform the Services in a professional and workmanlike '
        'manner, consistent with generally accepted industry standards and practices. The '
        'Service Provider represents that it has the skills, qualifications, and experience '
        'necessary to perform the Services.'
    )
    doc.add_paragraph(
        'Each party shall maintain the confidentiality of all Confidential Information received '
        'from the other party and shall not disclose such information to any third party without '
        'the prior written consent of the disclosing party, except as required by law or court order.'
    )
    doc.add_paragraph(
        'Neither party shall be liable to the other for any indirect, incidental, consequential, '
        'special, or exemplary damages arising out of or related to this Agreement, regardless '
        'of the theory of liability, even if such party has been advised of the possibility of '
        'such damages. The Service Provider\'s total aggregate liability under this Agreement '
        'shall not exceed the total fees paid by the Client.'
    )

    # --- Section 6: Intellectual Property ---
    doc.add_heading('Section 6: Intellectual Property', level=1)
    doc.add_paragraph(
        'All Deliverables created by the Service Provider specifically for the Client under '
        'this Agreement shall be considered "work made for hire" to the extent permitted by '
        'applicable law, and all intellectual property rights therein shall vest in the Client '
        'upon full payment of all fees due.'
    )
    doc.add_paragraph(
        'The Service Provider retains all rights in its pre-existing intellectual property, '
        'tools, methodologies, and frameworks ("Provider IP"). To the extent any Provider IP '
        'is incorporated into the Deliverables, the Service Provider hereby grants the Client '
        'a perpetual, non-exclusive, royalty-free license to use such Provider IP solely in '
        'connection with the Deliverables.'
    )

    # --- Section 7: Term and Termination ---
    doc.add_heading('Section 7: Term and Termination', level=1)
    doc.add_paragraph(
        'This Agreement shall commence on the Effective Date and continue for a period of '
        'eighteen (18) months, unless earlier terminated in accordance with this Section.'
    )
    doc.add_paragraph(
        'Either party may terminate this Agreement for cause upon thirty (30) days\' written '
        'notice if the other party materially breaches any provision of this Agreement and '
        'fails to cure such breach within the notice period.'
    )
    doc.add_paragraph(
        'The Client may terminate this Agreement for convenience upon sixty (60) days\' written '
        'notice, provided that the Client shall pay the Service Provider for all Services '
        'performed and expenses incurred through the effective date of termination, plus a '
        'termination fee equal to 15% of the remaining unpaid contract value.'
    )

    # --- Section 8: Insurance and Indemnification ---
    doc.add_heading('Section 8: Insurance and Indemnification', level=1)
    doc.add_paragraph(
        'The Service Provider shall maintain, at its own expense, throughout the term of this '
        'Agreement: (a) commercial general liability insurance with limits of not less than '
        '$2,000,000 per occurrence and $5,000,000 in the aggregate; (b) professional liability '
        '(errors and omissions) insurance with limits of not less than $3,000,000 per claim; '
        'and (c) workers\' compensation insurance as required by applicable law.'
    )
    doc.add_paragraph(
        'Each party shall indemnify, defend, and hold harmless the other party and its officers, '
        'directors, employees, and agents from and against any and all claims, damages, losses, '
        'liabilities, and expenses (including reasonable attorneys\' fees) arising out of or '
        'resulting from the indemnifying party\'s breach of this Agreement or negligent acts '
        'or omissions in connection with the performance of this Agreement.'
    )

    # --- Section 9: Governing Law and Dispute Resolution ---
    doc.add_heading('Section 9: Governing Law and Dispute Resolution', level=1)
    doc.add_paragraph(
        'This Agreement shall be governed by and construed in accordance with the laws of the '
        'State of New York, without regard to its conflicts of law principles. Any legal action '
        'or proceeding arising under this Agreement shall be brought exclusively in the federal '
        'or state courts located in the Borough of Manhattan, New York City, and the parties '
        'hereby irrevocably consent to the personal jurisdiction of such courts.'
    )
    doc.add_paragraph(
        'Prior to initiating any legal action, the parties shall attempt in good faith to resolve '
        'any dispute through mediation administered by the American Arbitration Association in '
        'accordance with its Commercial Mediation Procedures. If mediation is unsuccessful within '
        'sixty (60) days, either party may pursue its legal remedies.'
    )

    # --- Section 10: General Provisions ---
    doc.add_heading('Section 10: General Provisions', level=1)
    doc.add_paragraph(
        'This Agreement constitutes the entire agreement between the parties with respect to '
        'the subject matter hereof and supersedes all prior and contemporaneous agreements, '
        'understandings, negotiations, and discussions, whether oral or written.'
    )
    doc.add_paragraph(
        'This Agreement may not be amended or modified except by a written instrument signed '
        'by both parties. No waiver of any provision of this Agreement shall be effective unless '
        'in writing and signed by the party granting the waiver.'
    )
    doc.add_paragraph(
        'If any provision of this Agreement is held to be invalid or unenforceable, the remaining '
        'provisions shall continue in full force and effect. The invalidity or unenforceability '
        'of any provision shall not affect the validity or enforceability of any other provision.'
    )

    # --- Signature Block ---
    doc.add_paragraph('')
    sig = doc.add_paragraph()
    sig.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = sig.add_run('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.')
    run.font.size = Pt(11)

    doc.add_paragraph('')
    doc.add_paragraph('_______________________________          _______________________________')
    doc.add_paragraph('Meridian Technology Solutions, Inc.       Cascade Financial Group, LLC')
    doc.add_paragraph('By: Jennifer A. Whitfield, CEO            By: Robert K. Tanaka, Managing Partner')
    doc.add_paragraph('Date: _______________                      Date: _______________')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
