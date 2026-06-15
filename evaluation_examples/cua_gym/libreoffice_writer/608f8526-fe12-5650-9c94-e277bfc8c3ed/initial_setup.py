"""
Initial Setup: Create a Writer document with a manually-numbered outline
Task ID: writer_bs_057
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_057'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading('Quarterly Business Strategy Review', level=1)

    # Introductory paragraph
    doc.add_paragraph(
        'The following outline summarizes the key strategic initiatives, '
        'action items, and performance metrics discussed during the Q2 2025 '
        'quarterly review meeting held on June 12, 2025.'
    )

    doc.add_paragraph('')  # spacer

    # --- Manually numbered outline (NO list styles, just plain text) ---
    # Level 1 items - no indent, manual Roman numerals
    # Level 2 items - manual letters, typed indent via spaces
    # Level 3 items - manual Arabic numbers, typed indent via spaces

    items = [
        ("I. Revenue Growth Initiatives", 0),
        ("    A. Expand enterprise sales team by 15% in Q3", 1),
        ("    B. Launch partner referral program targeting mid-market segment", 1),
        ("        1. Identify top 25 potential channel partners by July 15", 2),
        ("        2. Develop co-marketing materials and joint webinar series", 2),
        ("        3. Establish revenue-sharing framework with tiered commissions", 2),
        ("    C. Increase average deal size through bundled service offerings", 1),
        ("II. Product Development Roadmap", 0),
        ("    A. Release version 3.2 with enhanced analytics dashboard", 1),
        ("        1. Complete user research interviews with 30 enterprise clients", 2),
        ("        2. Finalize wireframes and submit for design review by August 1", 2),
        ("    B. Integrate machine learning pipeline for predictive forecasting", 1),
        ("        1. Hire two senior ML engineers with time-series experience", 2),
        ("        2. Deploy initial prototype to staging environment by September", 2),
        ("        3. Conduct A/B testing with pilot customer group", 2),
        ("    C. Migrate legacy infrastructure to cloud-native architecture", 1),
        ("III. Customer Retention Strategy", 0),
        ("    A. Implement proactive health scoring for all accounts over $50K ARR", 1),
        ("        1. Define churn risk indicators based on usage telemetry data", 2),
        ("        2. Build automated alert system for customer success managers", 2),
        ("    B. Launch quarterly business review program for top 50 accounts", 1),
        ("    C. Reduce average support response time from 4 hours to under 2 hours", 1),
        ("        1. Add dedicated support tier for enterprise customers", 2),
        ("        2. Deploy AI-assisted ticket routing and suggested responses", 2),
        ("        3. Hire 8 additional support engineers across APAC and EMEA regions", 2),
        ("IV. Operational Excellence", 0),
        ("    A. Standardize project management methodology across all departments", 1),
        ("        1. Evaluate and select unified PM tooling by end of July", 2),
        ("        2. Conduct training workshops for all team leads in August", 2),
        ("    B. Reduce operational costs by 12% through process automation", 1),
        ("    C. Establish cross-functional tiger teams for critical initiatives", 1),
        ("        1. Appoint executive sponsors for each strategic priority", 2),
        ("        2. Define OKRs and reporting cadence for monthly check-ins", 2),
    ]

    for text, level in items:
        para = doc.add_paragraph(text)
        # All items use default paragraph style with no list formatting
        para.style = doc.styles['Normal']
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(2)

    # Closing paragraph
    doc.add_paragraph('')
    doc.add_paragraph(
        'Next review scheduled for September 18, 2025. All department heads '
        'should prepare progress updates against the targets outlined above.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
