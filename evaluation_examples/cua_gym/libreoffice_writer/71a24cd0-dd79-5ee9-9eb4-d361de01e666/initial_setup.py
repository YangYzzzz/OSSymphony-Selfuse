"""
Initial Setup: Standard legal brief with default formatting
Task ID: writer_legal_075
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_075'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Default page setup (standard Letter, 1-inch margins all around)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Document Title ---
    title = doc.add_heading('MEMORANDUM OF POINTS AND AUTHORITIES', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Case Caption ---
    caption = doc.add_paragraph()
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = caption.add_run('IN THE SUPERIOR COURT OF THE STATE OF CALIFORNIA')
    run.bold = True
    run.font.size = Pt(12)

    caption2 = doc.add_paragraph()
    caption2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = caption2.add_run('FOR THE COUNTY OF LOS ANGELES')
    run2.bold = True
    run2.font.size = Pt(12)

    doc.add_paragraph()  # blank line

    # --- Case Information ---
    case_info = doc.add_paragraph()
    run_c = case_info.add_run('PACIFIC COAST DEVELOPMENT, LLC,')
    run_c.bold = True
    case_info.add_run('\n\tPlaintiff,')

    vs = doc.add_paragraph()
    vs.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    vs.add_run('v.').bold = True

    case_info2 = doc.add_paragraph()
    run_d = case_info2.add_run('WESTSIDE CONSTRUCTION GROUP, INC.,')
    run_d.bold = True
    case_info2.add_run('\n\tDefendant.')

    case_num = doc.add_paragraph()
    case_num.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_cn = case_num.add_run('Case No. 24-CV-08391')
    run_cn.font.size = Pt(12)

    doc.add_paragraph()  # blank line

    # --- Section I ---
    h2 = doc.add_heading('I. INTRODUCTION', level=2)

    p1 = doc.add_paragraph()
    p1.paragraph_format.first_line_indent = Inches(0.5)
    p1.add_run(
        'This memorandum is submitted in support of Plaintiff Pacific Coast '
        'Development, LLC\'s Motion for Summary Judgment. As demonstrated below, '
        'there are no genuine issues of material fact in dispute, and Plaintiff '
        'is entitled to judgment as a matter of law pursuant to California Code '
        'of Civil Procedure Section 437c.'
    )

    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Inches(0.5)
    p2.add_run(
        'On or about March 15, 2023, Plaintiff and Defendant entered into a '
        'Construction Services Agreement (the "Agreement") whereby Defendant '
        'agreed to perform certain renovation and construction services at '
        'Plaintiff\'s commercial property located at 4521 Wilshire Boulevard, '
        'Los Angeles, California 90010 (the "Property").'
    )

    # --- Section II ---
    doc.add_heading('II. STATEMENT OF UNDISPUTED FACTS', level=2)

    p3 = doc.add_paragraph()
    p3.paragraph_format.first_line_indent = Inches(0.5)
    p3.add_run(
        '1. On March 15, 2023, Pacific Coast Development, LLC and Westside '
        'Construction Group, Inc. executed the Construction Services Agreement '
        '(Exhibit A) for renovations to the Property. The total contract price '
        'was $2,450,000.00, with a completion deadline of December 31, 2023.'
    )

    p4 = doc.add_paragraph()
    p4.paragraph_format.first_line_indent = Inches(0.5)
    p4.add_run(
        '2. Plaintiff made timely progress payments totaling $1,837,500.00 '
        '(representing 75% of the contract price) as documented in payment '
        'records (Exhibit B). Each payment was made within ten (10) business '
        'days of receipt of Defendant\'s invoices.'
    )

    p5 = doc.add_paragraph()
    p5.paragraph_format.first_line_indent = Inches(0.5)
    p5.add_run(
        '3. On November 8, 2023, Defendant ceased all work at the Property '
        'without prior notice. At the time of abandonment, independent '
        'inspection confirmed that only approximately 62% of the contracted '
        'work had been completed (Exhibit C, Inspector\'s Report).'
    )

    p6 = doc.add_paragraph()
    p6.paragraph_format.first_line_indent = Inches(0.5)
    p6.add_run(
        '4. Plaintiff provided written notice of default on November 20, 2023, '
        'as required by Section 8.2 of the Agreement, and allowed the '
        'contractual fifteen (15) day cure period. Defendant failed to resume '
        'work or respond to the notice of default (Exhibit D).'
    )

    p7 = doc.add_paragraph()
    p7.paragraph_format.first_line_indent = Inches(0.5)
    p7.add_run(
        '5. To complete the remaining work, Plaintiff retained Rivera & '
        'Associates General Contractors at a cost of $1,285,000.00, '
        'representing the completion costs for the unfinished 38% of the '
        'project scope (Exhibit E).'
    )

    # --- Section III ---
    doc.add_heading('III. ARGUMENT', level=2)

    doc.add_heading('A. Standard of Review', level=3)

    p8 = doc.add_paragraph()
    p8.paragraph_format.first_line_indent = Inches(0.5)
    p8.add_run(
        'Summary judgment is appropriate when "all the papers submitted show '
        'that there is no triable issue as to any material fact and that the '
        'moving party is entitled to a judgment as a matter of law." Cal. Code '
        'Civ. Proc. § 437c(c). The moving party bears the initial burden of '
        'production to make a prima facie showing of the nonexistence of any '
        'triable issue of material fact. Aguilar v. Atlantic Richfield Co. '
        '(2001) 25 Cal.4th 826, 850.'
    )

    doc.add_heading('B. Defendant Breached the Agreement', level=3)

    p9 = doc.add_paragraph()
    p9.paragraph_format.first_line_indent = Inches(0.5)
    p9.add_run(
        'The elements of a breach of contract claim under California law are: '
        '(1) the existence of a contract; (2) plaintiff\'s performance or '
        'excuse for nonperformance; (3) defendant\'s breach; and (4) resulting '
        'damages. Oasis West Realty, LLC v. Goldman (2011) 51 Cal.4th 811, 821.'
    )

    p10 = doc.add_paragraph()
    p10.paragraph_format.first_line_indent = Inches(0.5)
    p10.add_run(
        'Here, all four elements are established by undisputed evidence. The '
        'Agreement is authenticated (UF 1). Plaintiff performed by making '
        '$1,837,500.00 in progress payments (UF 2). Defendant breached by '
        'abandoning the project with only 62% completion (UF 3). And Plaintiff '
        'suffered damages of at least $672,500.00, calculated as the '
        'difference between the completion costs paid to Rivera & Associates '
        '($1,285,000.00) and the unpaid balance of the original contract '
        '($612,500.00).'
    )

    # --- Section IV ---
    doc.add_heading('IV. CONCLUSION', level=2)

    p11 = doc.add_paragraph()
    p11.paragraph_format.first_line_indent = Inches(0.5)
    p11.add_run(
        'For the foregoing reasons, Plaintiff Pacific Coast Development, LLC '
        'respectfully requests that this Court grant its Motion for Summary '
        'Judgment and enter judgment in its favor in the amount of $672,500.00, '
        'plus prejudgment interest at the legal rate from December 16, 2023, '
        'plus costs of suit and such other relief as the Court deems just and '
        'proper.'
    )

    doc.add_paragraph()  # blank line

    # --- Signature Block ---
    sig = doc.add_paragraph()
    sig.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sig.add_run('Dated: January 15, 2025')

    doc.add_paragraph()

    sig2 = doc.add_paragraph()
    sig2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    sig2.add_run('Respectfully submitted,\n\n\n')
    sig2.add_run('________________________________\n').font.size = Pt(12)
    run_name = sig2.add_run('Elena M. Vasquez, Esq.\n')
    run_name.bold = True
    sig2.add_run('VASQUEZ & CHEN LLP\n')
    sig2.add_run('350 South Grand Avenue, Suite 2800\n')
    sig2.add_run('Los Angeles, California 90071\n')
    sig2.add_run('Telephone: (213) 555-0142\n')
    sig2.add_run('Attorney for Plaintiff\n')
    sig2.add_run('Pacific Coast Development, LLC')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
