"""
Initial Setup: Set wrap spacing for text box on page 2
Task ID: writer_obj_044
Domain: libreoffice_writer

Creates tight_layout.docx with:
- Page 1: body text content
- Page 2: a 6cm x 4cm text box with Parallel (square) text wrapping
  and DEFAULT uniform wrap spacing (0.32cm all sides)
File location: /home/user/Desktop/tight_layout.docx
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_044'
OUTPUT = f'{WORKDIR}/tight_layout.docx'

# Default wrap spacing: 0.32cm ~ 115200 EMU (typical LibreOffice default)
DEFAULT_DIST = 115200  # ~0.32 cm in EMU


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_textbox_with_wrap(doc, text, width_cm, height_cm, dist_all_emu):
    """
    Add a floating text box (anchor shape) with square/parallel text wrapping
    and specified uniform wrap spacing (distT/B/L/R in EMU).
    """
    # We build the drawing XML manually for a floating text box with wrap
    # distT, distB, distL, distR: wrap distances in EMU
    dist = str(dist_all_emu)
    width_emu = str(int(Cm(width_cm)))
    height_emu = str(int(Cm(height_cm)))

    drawing_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:r>
    <w:rPr/>
    <w:drawing>
      <wp:anchor
        xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        distT="{dist}" distB="{dist}" distL="{dist}" distR="{dist}"
        simplePos="0" relativeHeight="251658240" behindDoc="0"
        locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column">
          <wp:posOffset>914400</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="paragraph">
          <wp:posOffset>457200</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapSquare wrapText="bothSides"/>
        <wp:docPr id="1" name="TextBox 1" descr="Text box with parallel wrapping"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
              <wps:cNvSpPr txBx="1">
                <a:spLocks noChangeArrowheads="1"/>
              </wps:cNvSpPr>
              <wps:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{width_emu}" cy="{height_emu}"/>
                </a:xfrm>
                <a:prstGeom prst="rect">
                  <a:avLst/>
                </a:prstGeom>
                <a:solidFill>
                  <a:srgbClr val="FFFFFF"/>
                </a:solidFill>
                <a:ln>
                  <a:solidFill>
                    <a:srgbClr val="000000"/>
                  </a:solidFill>
                </a:ln>
              </wps:spPr>
              <wps:txbx>
                <w:txbxContent>
                  <w:p>
                    <w:r>
                      <w:t>{text}</w:t>
                    </w:r>
                  </w:p>
                </w:txbxContent>
              </wps:txbx>
              <wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="overflow"
                          horzOverflow="overflow" vert="horz" wrap="square"
                          lIns="91440" tIns="45720" rIns="91440" bIns="45720"
                          numCol="1" spcCol="0" rtlCol="0" fromWordArt="0"
                          anchor="t" anchorCtr="0" forceAA="0" compatLnSpc="1">
                <a:prstTxWarp prst="textNoShape">
                  <a:avLst/>
                </a:prstTxWarp>
              </wps:bodyPr>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </w:r>
</w:p>'''

    p_elem = etree.fromstring(drawing_xml)
    doc.add_paragraph()  # placeholder, we'll replace last para
    # Replace the last added paragraph element
    doc.paragraphs[-1]._element.getparent().replace(
        doc.paragraphs[-1]._element, p_elem
    )


def create_initial():
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    # --- Page 1: Content ---
    heading = doc.add_heading('Document Layout Guide', level=1)

    doc.add_paragraph(
        'This document provides an overview of effective text layout strategies for '
        'professional reports and publications. Proper use of text boxes, wrapping '
        'styles, and spacing ensures a clean, readable appearance throughout.'
    )

    doc.add_heading('Section 1: Layout Principles', level=2)

    doc.add_paragraph(
        'Modern document design relies on clear visual hierarchy. Headers, subheadings, '
        'and body text each play a distinct role. Consistent margins and spacing between '
        'elements create a sense of order and professionalism.'
    )

    doc.add_paragraph(
        'When embedding figures or call-out boxes within text, text wrapping determines '
        'how surrounding paragraphs flow around the object. Common wrapping modes include '
        'Square, Tight, Through, Top and Bottom (Parallel), and In Line with Text.'
    )

    doc.add_heading('Section 2: Text Wrapping Modes', level=2)

    doc.add_paragraph(
        'Square wrapping allows text to flow in a rectangular boundary around the object. '
        'This is often used for images and pull quotes where a clean rectangular margin '
        'is desired. The wrap spacing controls the distance between the text and the '
        'edges of the object.'
    )

    doc.add_paragraph(
        'Parallel (Top and Bottom) wrapping places the object on its own horizontal band, '
        'with text appearing above and below but not to the sides. This mode is suitable '
        'for wide objects that span most of the column width.'
    )

    doc.add_paragraph(
        'The wrap spacing values—top, bottom, left, and right—define the clearance between '
        'the wrapped object and adjacent text. Adjusting these values fine-tunes the visual '
        'balance of the layout.'
    )

    # Page break to page 2
    doc.add_page_break()

    # --- Page 2: Content with text box ---
    doc.add_heading('Section 3: Advanced Layout Techniques', level=2)

    doc.add_paragraph(
        'On this page we demonstrate a text box with Parallel (Top and Bottom) wrapping. '
        'The box below is positioned within the text flow, and the surrounding paragraphs '
        'appear above and below it. The current wrap spacing uses the default uniform value '
        'on all four sides.'
    )

    # Add the text box with default uniform spacing
    add_textbox_with_wrap(
        doc,
        text='Key Insight: Parallel wrapping keeps the text box centered in the text flow.',
        width_cm=6,
        height_cm=4,
        dist_all_emu=DEFAULT_DIST,
    )

    doc.add_paragraph(
        'After the text box, the document continues with additional content. Notice how '
        'the paragraph text resumes directly below the text box boundary, respecting the '
        'wrap spacing that separates the box from the surrounding text.'
    )

    doc.add_paragraph(
        'Designers often reduce top and bottom spacing to zero when they want the text '
        'to appear immediately adjacent to the box, while keeping side margins to prevent '
        'the text from crowding the box horizontally.'
    )

    doc.add_paragraph(
        'This technique is especially effective in newsletter-style layouts where space '
        'is at a premium and the designer wants to maximize text density while still '
        'maintaining clear visual separation on the sides.'
    )

    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
