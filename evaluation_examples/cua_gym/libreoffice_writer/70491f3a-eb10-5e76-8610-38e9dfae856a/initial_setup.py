"""
Initial Setup: Masters thesis document without a Table of Contents
Task ID: writer_struct_028
Domain: libreoffice_writer

Creates a 15-page master's thesis document with 8 Heading 1 entries
but NO table of contents. The agent must insert a TOC with title 'Contents'.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_028'
# Task context says the file is at ~/Desktop/masters_thesis.docx
OUTPUT = f'{WORKDIR}/Desktop/masters_thesis.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # --- Title Page ---
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    run = title_para.add_run('The Impact of Machine Learning on Climate Change Prediction Models')
    run.bold = True
    run.font.size = Pt(16)

    sub_para = doc.add_paragraph()
    sub_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_para.paragraph_format.space_before = Pt(24)
    run2 = sub_para.add_run('A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nMaster of Science in Environmental Informatics')
    run2.font.size = Pt(12)

    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.space_before = Pt(36)
    run3 = author_para.add_run('By\nEmily Hartwell\nDepartment of Environmental Science and Informatics\nUniversity of Greenfield\n2024')
    run3.font.size = Pt(12)

    doc.add_page_break()

    # --- Acknowledgements ---
    doc.add_heading('Acknowledgements', level=1)
    doc.add_paragraph(
        'I would like to express my deepest gratitude to my thesis advisor, Professor '
        'Dr. Samuel Reinhart, whose invaluable guidance and unwavering support made this '
        'research possible. His expertise in environmental informatics and machine learning '
        'has been a constant source of inspiration throughout this journey.'
    )
    doc.add_paragraph(
        'I am also deeply grateful to the members of my thesis committee: Professor '
        'Dr. Adaeze Okonkwo, whose insights into climate modeling greatly enriched my work, '
        'and Professor Dr. Lars Bergstrom, whose thorough review of statistical methods '
        'helped me refine my analytical approach. Their constructive feedback during each '
        'committee meeting was invaluable.'
    )
    doc.add_paragraph(
        'Special thanks go to my colleagues at the Environmental Data Lab — particularly '
        'Marcus Fielding, Priya Nair, and Tobias Schreiber — for their camaraderie, '
        'stimulating discussions, and technical support. Our collaborative problem-solving '
        'sessions helped me navigate many challenging aspects of this research.'
    )
    doc.add_paragraph(
        'I am grateful to the University of Greenfield\'s High Performance Computing Center '
        'for providing the computational resources necessary for running the large-scale '
        'climate simulations. Without access to their infrastructure, this project would '
        'not have been feasible within the given timeframe.'
    )
    doc.add_paragraph(
        'Finally, I owe an immeasurable debt of gratitude to my family — my parents, '
        'David and Catherine Hartwell, and my sister, Sophie — for their unconditional love '
        'and constant encouragement. Their belief in my abilities, even during the most '
        'challenging periods of this work, provided the motivation I needed to persevere.'
    )

    doc.add_page_break()

    # --- Introduction ---
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'Climate change represents one of the most complex and consequential challenges '
        'facing humanity in the twenty-first century. The ability to accurately predict '
        'future climate patterns is essential for developing effective mitigation and '
        'adaptation strategies. Traditional physics-based climate models, while sophisticated, '
        'face limitations in computational efficiency and the ability to capture fine-scale '
        'regional variations.'
    )
    doc.add_paragraph(
        'The emergence of machine learning (ML) techniques has created new opportunities '
        'for enhancing climate prediction capabilities. By leveraging large historical '
        'datasets and pattern recognition algorithms, ML models can potentially identify '
        'complex nonlinear relationships that are difficult to represent in traditional '
        'numerical models. Recent advances in deep learning, in particular, have shown '
        'promise in various meteorological applications.'
    )
    doc.add_paragraph(
        'This thesis investigates the integration of machine learning algorithms with '
        'established climate modeling frameworks to improve the accuracy and resolution '
        'of regional climate predictions. Specifically, we explore the application of '
        'Long Short-Term Memory (LSTM) networks and Convolutional Neural Networks (CNNs) '
        'in conjunction with the Community Earth System Model (CESM) to enhance predictions '
        'of temperature anomalies, precipitation patterns, and extreme weather events '
        'across the North Atlantic region for the period 2025-2075.'
    )

    heading_intro_sub = doc.add_heading('Research Objectives', level=2)
    doc.add_paragraph(
        'The primary objectives of this research are: (1) to evaluate the performance of '
        'selected machine learning algorithms in climate prediction tasks against established '
        'baselines; (2) to develop a hybrid modeling framework that leverages the strengths '
        'of both physics-based and data-driven approaches; (3) to assess the uncertainty '
        'quantification capabilities of ML-enhanced climate models; and (4) to provide '
        'actionable insights for policymakers regarding projected climate scenarios.'
    )

    doc.add_heading('Scope and Limitations', level=2)
    doc.add_paragraph(
        'The scope of this research is limited to regional climate predictions for the '
        'North Atlantic basin. While the methodologies developed may be applicable to '
        'other regions, generalizability has not been explicitly tested. Additionally, '
        'this work focuses on atmospheric variables and does not incorporate oceanic '
        'circulation models beyond their boundary condition roles.'
    )

    doc.add_page_break()

    # --- Literature Review ---
    doc.add_heading('Literature Review', level=1)
    doc.add_paragraph(
        'The application of statistical and machine learning methods to climate science '
        'has a rich history dating back several decades. Early work by Huth (1999) and '
        'Zorita and von Storch (1999) demonstrated the utility of statistical downscaling '
        'techniques using regression-based approaches. These methods established a foundation '
        'for subsequent data-driven climate modeling efforts.'
    )
    doc.add_paragraph(
        'The rise of neural networks in the 1990s introduced new possibilities for climate '
        'modeling. Tangang et al. (1998) pioneered the use of artificial neural networks '
        'for sea surface temperature prediction, demonstrating superior performance to '
        'linear statistical models in certain conditions. However, the computational '
        'constraints of the era limited the complexity and scale of these applications.'
    )
    doc.add_paragraph(
        'The deep learning revolution, catalyzed by breakthroughs in image recognition '
        '(Krizhevsky et al., 2012) and natural language processing (Vaswani et al., 2017), '
        'has profoundly influenced climate science. Reichstein et al. (2019) provided an '
        'influential overview of deep learning applications in Earth system sciences, '
        'highlighting both opportunities and challenges. Rasp et al. (2018) demonstrated '
        'that deep neural networks can emulate atmospheric convection parameterizations '
        'with high accuracy, opening new avenues for improving climate model efficiency.'
    )

    doc.add_heading('Machine Learning in Climate Prediction', level=2)
    doc.add_paragraph(
        'Recurrent neural networks, particularly LSTM architectures, have demonstrated '
        'particular promise for time-series climate data. Shi et al. (2015) introduced '
        'the ConvLSTM architecture for precipitation nowcasting, combining spatial and '
        'temporal learning capabilities. Ham et al. (2019) achieved state-of-the-art '
        'ENSO prediction performance using transfer learning with CNNs, significantly '
        'outperforming dynamical models at lead times beyond six months.'
    )

    doc.add_heading('Hybrid Modeling Approaches', level=2)
    doc.add_paragraph(
        'The integration of physics-based and machine learning approaches represents a '
        'promising frontier in climate modeling. Schneider et al. (2017) proposed using '
        'ML to improve climate model parameterizations, while Brenowitz and Bretherton '
        '(2018) demonstrated the feasibility of replacing convective parameterizations '
        'with neural networks. More recently, Pathak et al. (2022) introduced FourCastNet, '
        'a Fourier-based neural network model capable of generating global medium-range '
        'weather forecasts at a fraction of the computational cost of traditional NWP models.'
    )

    doc.add_page_break()

    # --- Methodology ---
    doc.add_heading('Methodology', level=1)
    doc.add_paragraph(
        'This research employs a mixed-methods approach combining quantitative climate '
        'data analysis with computational modeling. The methodology is structured in three '
        'phases: data collection and preprocessing, model development and training, and '
        'validation and uncertainty analysis.'
    )

    doc.add_heading('Data Sources and Preprocessing', level=2)
    doc.add_paragraph(
        'Historical climate data was obtained from the ERA5 reanalysis dataset (Hersbach '
        'et al., 2020), provided by the European Centre for Medium-Range Weather Forecasts '
        '(ECMWF). The dataset covers the period 1940-2023 at a spatial resolution of '
        '0.25° × 0.25° and hourly temporal resolution. Key variables extracted include '
        '2-meter air temperature, total precipitation, sea level pressure, 500 hPa '
        'geopotential height, and sea surface temperature.'
    )
    doc.add_paragraph(
        'Data preprocessing involved several steps: quality control and outlier removal '
        'using a 3-sigma criterion; temporal aggregation to daily, monthly, and seasonal '
        'means; spatial regridding to a uniform 1° × 1° grid using bilinear interpolation; '
        'and normalization using z-score standardization calculated over the 1981-2010 '
        'reference period.'
    )

    doc.add_heading('Model Architecture', level=2)
    doc.add_paragraph(
        'Two primary model architectures were developed and evaluated: (1) a Bidirectional '
        'LSTM network with attention mechanisms for temporal sequence prediction, and (2) '
        'a Convolutional LSTM (ConvLSTM) for spatiotemporal pattern recognition. Both '
        'architectures were implemented using TensorFlow 2.10 and trained on the University '
        'of Greenfield\'s High Performance Computing cluster using NVIDIA A100 GPUs.'
    )

    doc.add_page_break()

    # --- Results ---
    doc.add_heading('Results', level=1)
    doc.add_paragraph(
        'This section presents the key findings from the evaluation of ML-enhanced climate '
        'prediction models. Results are organized by prediction task and temporal horizon, '
        'with comparisons to baseline dynamical models and climatological benchmarks.'
    )

    doc.add_heading('Temperature Anomaly Prediction', level=2)
    doc.add_paragraph(
        'The Bidirectional LSTM model achieved a Root Mean Square Error (RMSE) of 0.42°C '
        'for monthly mean temperature anomaly prediction over the North Atlantic region '
        'at a lead time of 3 months, representing a 23% improvement over the CESM baseline '
        '(RMSE = 0.55°C). Performance degraded gracefully with increasing lead time, with '
        'RMSE values of 0.61°C and 0.84°C at 6-month and 12-month lead times respectively.'
    )
    doc.add_paragraph(
        'Spatial analysis of prediction skill revealed that the ML model performed '
        'particularly well over the Gulf Stream extension region, where SST gradients '
        'provide strong predictive signals. Performance was comparatively lower over '
        'high-latitude regions north of 60°N, likely due to increased internal variability '
        'and reduced training data coverage for extreme conditions.'
    )

    doc.add_heading('Precipitation Pattern Analysis', level=2)
    doc.add_paragraph(
        'Precipitation prediction proved more challenging than temperature, consistent '
        'with the known high spatial and temporal variability of rainfall patterns. The '
        'ConvLSTM model achieved a Heidke Skill Score (HSS) of 0.38 for monthly precipitation '
        'anomaly prediction at a 1-month lead time, compared to 0.29 for the CESM baseline. '
        'However, the skill decreased rapidly with lead time, with HSS values near zero '
        'at 6-month lead times for all models tested.'
    )

    doc.add_page_break()

    # --- Discussion ---
    doc.add_heading('Discussion', level=1)
    doc.add_paragraph(
        'The results of this study demonstrate that machine learning approaches can '
        'meaningfully enhance climate prediction capabilities, particularly for temperature '
        'anomalies at seasonal timescales. The 23% improvement in temperature prediction '
        'skill over the dynamical baseline is comparable to or exceeds improvements reported '
        'in similar studies (Ham et al., 2019; Weyn et al., 2020), suggesting that the '
        'hybrid LSTM-CESM framework developed here represents a viable approach for '
        'operational seasonal forecasting.'
    )
    doc.add_paragraph(
        'The relative underperformance of ML models for precipitation prediction highlights '
        'a fundamental challenge in climate modeling: precipitation is inherently more '
        'difficult to predict due to its dependence on convective processes that operate '
        'at scales below the resolution of most climate models. The modest skill improvement '
        'achieved at 1-month lead times suggests that ML can capture some predictable '
        'large-scale circulation patterns associated with precipitation anomalies, but '
        'cannot overcome the fundamental limits imposed by atmospheric chaos at longer '
        'lead times.'
    )

    doc.add_heading('Implications for Climate Policy', level=2)
    doc.add_paragraph(
        'The improved temperature prediction skill demonstrated in this study has practical '
        'implications for climate adaptation planning. More accurate seasonal temperature '
        'forecasts could benefit agricultural planning, energy demand forecasting, and '
        'public health preparedness for heat-related illnesses. However, the uncertainty '
        'quantification results suggest that careful communication of forecast uncertainty '
        'will be essential to avoid overconfidence in model projections.'
    )

    doc.add_heading('Limitations and Future Work', level=2)
    doc.add_paragraph(
        'Several limitations of this study should be acknowledged. First, the model was '
        'trained and evaluated exclusively on ERA5 reanalysis data, which itself has '
        'inherent uncertainties and may not perfectly represent the true atmospheric state. '
        'Future work should evaluate model performance against independent observational '
        'datasets. Second, the current framework does not incorporate socioeconomic '
        'scenarios or emissions trajectories, which are essential for long-term climate '
        'projections. Integrating Representative Concentration Pathway (RCP) or Shared '
        'Socioeconomic Pathway (SSP) scenarios into the ML framework represents an '
        'important direction for future research.'
    )

    doc.add_page_break()

    # --- Conclusion ---
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'This thesis has investigated the integration of machine learning algorithms — '
        'specifically Bidirectional LSTM networks and ConvLSTM architectures — with the '
        'Community Earth System Model to enhance regional climate predictions over the '
        'North Atlantic basin. The study demonstrates that ML-enhanced hybrid modeling '
        'approaches can achieve meaningful improvements in predictive skill for temperature '
        'anomalies at seasonal timescales, while the improvement for precipitation '
        'prediction remains modest.'
    )
    doc.add_paragraph(
        'The key contributions of this research are: (1) the development of a novel '
        'hybrid ML-dynamical modeling framework that leverages the complementary strengths '
        'of physics-based and data-driven approaches; (2) a comprehensive evaluation of '
        'ML prediction skill across multiple variables, lead times, and spatial domains; '
        '(3) a rigorous uncertainty quantification methodology adapted for ML-enhanced '
        'climate predictions; and (4) actionable recommendations for integrating ML '
        'forecasts into operational climate services.'
    )
    doc.add_paragraph(
        'The findings support a cautiously optimistic view of the potential for machine '
        'learning to contribute to climate science. Rather than replacing traditional '
        'physics-based models, ML approaches appear most valuable as complementary tools '
        'that can improve efficiency, identify patterns in large datasets, and enhance '
        'specific aspects of model performance. The hybrid framework developed in this '
        'thesis represents one step toward a future where ML and dynamical modeling '
        'work in concert to advance our understanding and prediction of Earth\'s climate system.'
    )

    doc.add_page_break()

    # --- References ---
    doc.add_heading('References', level=1)

    references = [
        'Brenowitz, N. D., & Bretherton, C. S. (2018). Prognostic validation of a neural network unified physics parameterization. Geophysical Research Letters, 45(12), 6289-6298.',
        'Ham, Y. G., Kim, J. H., & Luo, J. J. (2019). Deep learning for multi-year ENSO forecasts. Nature, 573(7775), 568-572.',
        'Hersbach, H., Bell, B., Berrisford, P., Hirahara, S., Horányi, A., Muñoz-Sabater, J., ... & Thépaut, J. N. (2020). The ERA5 global reanalysis. Quarterly Journal of the Royal Meteorological Society, 146(730), 1999-2049.',
        'Huth, R. (1999). Statistical downscaling in central Europe: Evaluation of methods and potential predictors. Climate Research, 13(2), 91-101.',
        'Krizhevsky, A., Sutskever, I., & Hinton, G. E. (2012). ImageNet classification with deep convolutional neural networks. Advances in Neural Information Processing Systems, 25.',
        'Pathak, J., Subramanian, S., Harrington, P., Raja, S., Chattopadhyay, A., Mardani, M., ... & Kashinath, K. (2022). FourCastNet: A global data-driven high-resolution weather model using adaptive Fourier neural operators. arXiv preprint arXiv:2202.11214.',
        'Rasp, S., Pritchard, M. S., & Gentine, P. (2018). Deep learning to represent subgrid processes in climate models. Proceedings of the National Academy of Sciences, 115(39), 9684-9689.',
        'Reichstein, M., Camps-Valls, G., Stevens, B., Jung, M., Denzler, J., & Carvalhais, N. (2019). Deep learning and process understanding for data-driven Earth system science. Nature, 566(7743), 195-204.',
        'Schneider, T., Lan, S., Stuart, A., & Teixeira, J. (2017). Earth system modeling 2.0: A blueprint for models that learn from observations and targeted high-resolution simulations. Geophysical Research Letters, 44(24), 12-396.',
        'Shi, X., Chen, Z., Wang, H., Yeung, D. Y., Wong, W. K., & Woo, W. C. (2015). Convolutional LSTM network: A machine learning approach for precipitation nowcasting. Advances in Neural Information Processing Systems, 28.',
        'Tangang, F. T., Tang, B., Monahan, A. H., & Hsieh, W. W. (1998). Forecasting ENSO events: A neural network–extended EOF approach. Journal of Climate, 11(1), 29-41.',
        'Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., ... & Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30.',
        'Weyn, J. A., Durran, D. R., & Caruana, R. (2020). Improving data-driven global weather prediction using deep convolutional neural networks on a cubed sphere. Journal of Advances in Modeling Earth Systems, 12(9), e2020MS002109.',
        'Zorita, E., & von Storch, H. (1999). The analog method as a simple statistical downscaling technique: Comparison with more complicated methods. Journal of Climate, 12(8), 2474-2489.',
    ]

    for ref in references:
        ref_para = doc.add_paragraph(ref, style='Normal')
        ref_para.paragraph_format.left_indent = Inches(0.5)
        ref_para.paragraph_format.first_line_indent = Inches(-0.5)
        ref_para.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
