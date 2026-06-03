"""
Reward Script: Insert a table of contents with title 'Contents' at the beginning of the thesis.
Task ID: writer_struct_028
Domain: libreoffice_writer
Scoring:
  Component 1: A 'TOC Heading' style paragraph exists in the document (0.4 pts)
  Component 2: The TOC heading text is exactly 'Contents' (not 'Table of Contents') (0.4 pts)
  Component 3: The TOC heading appears before the first Heading 1 in the document (0.2 pts)
"""

import os
from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_struct_028'
FILE_PATH = f'{WORKDIR}/masters_thesis.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Task: Insert a TOC at the beginning of the thesis with title 'Contents'.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: A paragraph with 'TOC Heading' style exists (0.4 points)
    # This verifies that a TOC heading element was inserted.
    # In the initial document there is no 'TOC Heading' paragraph.
    try:
        toc_heading_para = None
        toc_heading_index = None
        for i, para in enumerate(doc.paragraphs):
            if para.style.name == 'TOC Heading':
                toc_heading_para = para
                toc_heading_index = i
                break  # we only care about the first one

        if toc_heading_para is not None:
            print(f"PASS: Component 1 — 'TOC Heading' style paragraph found at index {toc_heading_index} (0.4 pts)")
            total_score += 0.4
        else:
            print("FAIL: Component 1 — No paragraph with 'TOC Heading' style found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: The TOC heading text is exactly 'Contents' (0.4 points)
    # The task specifically requires 'Contents' instead of the default 'Table of Contents'.
    # Only award points if a TOC heading exists AND the text is exactly 'Contents'.
    try:
        if toc_heading_para is not None:
            toc_title_text = toc_heading_para.text.strip()
            if toc_title_text == 'Contents':
                print(f"PASS: Component 2 — TOC title is exactly 'Contents' (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 2 — TOC title is {repr(toc_title_text)}, expected exactly 'Contents'")
        else:
            print("FAIL: Component 2 — No TOC heading found, cannot check title text")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: The TOC heading appears before the first Heading 1 (0.2 points)
    # The TOC should be at the beginning of the document, before body sections begin.
    try:
        first_h1_index = None
        for i, para in enumerate(doc.paragraphs):
            if para.style.name == 'Heading 1':
                first_h1_index = i
                break

        if toc_heading_para is not None and first_h1_index is not None:
            if toc_heading_index < first_h1_index:
                print(f"PASS: Component 3 — TOC heading (index {toc_heading_index}) is before first Heading 1 (index {first_h1_index}) (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — TOC heading (index {toc_heading_index}) is NOT before first Heading 1 (index {first_h1_index})")
        elif toc_heading_para is None:
            print("FAIL: Component 3 — No TOC heading found, cannot check position")
        elif first_h1_index is None:
            print("FAIL: Component 3 — No Heading 1 found in document")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
