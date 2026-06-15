"""
Initial Setup: Resume document with inconsistent paragraph spacing
Task ID: writer_para_040
Domain: libreoffice_writer

Creates a resume DOCX with 11 paragraphs where:
- Section headings (Professional Experience, Education, Skills) have space_before=0pt
- Content paragraphs have intentionally inconsistent spacing (NOT 0pt before / 6pt after)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_para_040'
# File path as referenced in task context
OUTPUT = f'{WORKDIR}/Desktop/resume_draft.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Remove default style spacing so we control it explicitly
    # We use Normal style paragraphs for content

    # --- Paragraph 1: Name (center-aligned, bold) ---
    p1 = doc.add_paragraph()
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run1 = p1.add_run('JENNIFER MARTINEZ')
    run1.bold = True
    # Inconsistent spacing: some non-zero space_before, no space_after
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)

    # --- Paragraph 2: Contact info (center-aligned) ---
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p2.add_run('jennifer.martinez@email.com | (555) 123-4567 | LinkedIn: /in/jmartinez')
    # Inconsistent: 3pt before, 10pt after
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after = Pt(10)

    # --- Paragraph 3: Section heading "Professional Experience" (bold, space_before=0pt) ---
    p3 = doc.add_paragraph()
    run3 = p3.add_run('Professional Experience')
    run3.bold = True
    # Task context confirms space_before=0pt currently for section headings
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(4)

    # --- Paragraph 4: Job title (inconsistent spacing) ---
    p4 = doc.add_paragraph()
    p4.add_run('Senior Software Engineer \u2014 Amazon Web Services, Seattle, WA (2021-Present)')
    p4.paragraph_format.space_before = Pt(6)
    p4.paragraph_format.space_after = Pt(0)

    # --- Paragraph 5: Bullet 1 (inconsistent spacing) ---
    p5 = doc.add_paragraph()
    p5.add_run('Led a team of 8 engineers in developing a real-time data processing pipeline handling 2 million events per second.')
    p5.paragraph_format.space_before = Pt(0)
    p5.paragraph_format.space_after = Pt(3)

    # --- Paragraph 6: Bullet 2 (inconsistent spacing) ---
    p6 = doc.add_paragraph()
    p6.add_run('Designed and implemented a microservices architecture that reduced system latency by 45%.')
    p6.paragraph_format.space_before = Pt(2)
    p6.paragraph_format.space_after = Pt(8)

    # --- Paragraph 7: Section heading "Education" (bold, space_before=0pt) ---
    p7 = doc.add_paragraph()
    run7 = p7.add_run('Education')
    run7.bold = True
    # Task context confirms space_before=0pt currently
    p7.paragraph_format.space_before = Pt(0)
    p7.paragraph_format.space_after = Pt(4)

    # --- Paragraph 8: Degree 1 (inconsistent spacing) ---
    p8 = doc.add_paragraph()
    p8.add_run('Master of Science in Computer Science \u2014 Stanford University (2019-2021)')
    p8.paragraph_format.space_before = Pt(4)
    p8.paragraph_format.space_after = Pt(0)

    # --- Paragraph 9: Degree 2 (inconsistent spacing) ---
    p9 = doc.add_paragraph()
    p9.add_run('Bachelor of Science in Mathematics \u2014 UC Berkeley (2015-2019)')
    p9.paragraph_format.space_before = Pt(0)
    p9.paragraph_format.space_after = Pt(5)

    # --- Paragraph 10: Section heading "Skills" (bold, space_before=0pt) ---
    p10 = doc.add_paragraph()
    run10 = p10.add_run('Skills')
    run10.bold = True
    # Task context confirms space_before=0pt currently
    p10.paragraph_format.space_before = Pt(0)
    p10.paragraph_format.space_after = Pt(4)

    # --- Paragraph 11: Skills list (inconsistent spacing) ---
    p11 = doc.add_paragraph()
    p11.add_run('Python, Java, Go, AWS, Kubernetes, Terraform, PostgreSQL, MongoDB, Apache Kafka')
    p11.paragraph_format.space_before = Pt(1)
    p11.paragraph_format.space_after = Pt(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
