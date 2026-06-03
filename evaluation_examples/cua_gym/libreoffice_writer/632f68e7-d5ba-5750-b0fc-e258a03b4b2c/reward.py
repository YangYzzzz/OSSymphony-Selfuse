"""
FINAL REWARD SCRIPT - SUCCESS
Task: Align the project title to the center of the page width.
Generated: 2025-10-14 11:04:11
Status: success
Model: azure-o3
Total Steps: 3
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ------------------ Reward Script ------------------

def verify_project_title_centered(file_path: str) -> float:
    """Verify that the project title (assumed to be the first non-empty paragraph)
    is centred across the page width.

    Progressive scoring:
        +0.7 – First non-empty paragraph is explicitly centred.
        +0.3 – At least one centred, non-empty paragraph exists anywhere in the document.
    The score is capped at 1.0. No points are awarded for natural conditions such as
    file existence or successful loading (these are prerequisites)."""

    print(f"Checking file: {file_path}\n")
    total_score = 0.0
    max_score = 1.0

    # Prerequisite: file must exist and be readable (0 points)
    if not os.path.exists(file_path):
        print("✗ File not found – cannot verify.")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Failed to open document: {e}")
        return 0.0

    # Gather all non-empty paragraphs
    text_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    if not text_paragraphs:
        print("✗ No textual paragraphs found – nothing to verify.")
        return 0.0

    # Core requirement: first textual paragraph is centred
    first_para = text_paragraphs[0]
    alignment = first_para.paragraph_format.alignment  # May be None
    print(f"First paragraph text: '{first_para.text.strip()}'")
    print(f"First paragraph alignment value: {alignment}")

    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        print("✓ First textual paragraph is centred (0.7 pts)")
        total_score += 0.7
    else:
        print("✗ First textual paragraph is NOT centred (0 pts)")

    # Partial credit: any centred, non-empty paragraph exists
    any_centered = any(
        p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER and p.text.strip()
        for p in doc.paragraphs
    )

    if any_centered:
        print("✓ Found at least one centred paragraph in document (0.3 pts)")
        total_score += 0.3
    else:
        print("✗ No centred paragraphs found in entire document (0 pts)")

    # Cap at maximum
    final_score = min(total_score, max_score)

    print(f"\nTotal score: {final_score} (out of {max_score})")
    print(f"REWARD: {final_score}")
    return final_score

# ------------------ Execute Verification ------------------

if __name__ == "__main__":
    DOC_PATH = "/home/user/align_the_project_title_to_the_center_of_the_page_width.docx"
    verify_project_title_centered(DOC_PATH)

