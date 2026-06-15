"""
Reward Script: Create a version history table at the beginning of a technical specification
Task ID: writer_tech_028
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table exists with 4 columns
  Component 2 (0.25): Header row has correct column names (Version, Date, Author, Changes)
  Component 3 (0.30): Table has exactly 3 data rows with non-empty content
  Component 4 (0.20): Table is positioned before the main content (before "1. Introduction")
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_028'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with 4 columns (0.25 points)
    # Initial env has 0 tables, golden has 1 table with 4 cols
    target_table = None
    try:
        tables = doc.tables
        if len(tables) >= 1:
            # Find a table with exactly 4 columns
            for t in tables:
                if len(t.columns) == 4:
                    target_table = t
                    break
            if target_table is not None:
                print(f"PASS: Component 1 -- Found table with 4 columns (0.25 pts)")
                total_score += 0.25
            else:
                col_counts = [len(t.columns) for t in tables]
                print(f"FAIL: Component 1 -- No table with 4 columns found. Column counts: {col_counts}")
        else:
            print(f"FAIL: Component 1 -- No tables found in document (found {len(tables)})")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Header row has correct column names (0.25 points)
    # Expected: Version, Date, Author, Changes
    try:
        if target_table is not None:
            header_cells = [cell.text.strip().lower() for cell in target_table.rows[0].cells]
            expected_headers = ['version', 'date', 'author', 'changes']
            if header_cells == expected_headers:
                print(f"PASS: Component 2 -- Header row matches: {header_cells} (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 -- Header row mismatch. Expected {expected_headers}, found {header_cells}")
        else:
            print(f"FAIL: Component 2 -- No suitable 4-column table found to check headers")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Table has exactly 3 data rows with non-empty content (0.30 points)
    # Initial env has no table at all, so this only passes on golden
    try:
        if target_table is not None:
            data_rows = list(target_table.rows)[1:]  # Skip header
            num_data_rows = len(data_rows)
            if num_data_rows == 3:
                # Check that each data row has non-empty content in all cells
                empty_cells = []
                for ri, row in enumerate(data_rows):
                    for ci, cell in enumerate(row.cells):
                        if not cell.text.strip():
                            empty_cells.append((ri + 1, ci))
                            print(f"  WARN: Row {ri+1}, Col {ci} is empty")
                if len(empty_cells) == 0:
                    print(f"PASS: Component 3 -- 3 data rows, all cells populated (0.30 pts)")
                    total_score += 0.30
                else:
                    print(f"FAIL: Component 3 -- 3 data rows found but some cells are empty")
            else:
                print(f"FAIL: Component 3 -- Expected 3 data rows, found {num_data_rows}")
        else:
            print(f"FAIL: Component 3 -- No suitable 4-column table found to check data rows")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Table is positioned before "1. Introduction" heading (0.20 points)
    # The task says "at the beginning of the document". In the golden, the table
    # appears after the title area and before section 1. In initial, there is no table.
    try:
        if target_table is not None:
            # Find the XML position of the table vs the "1. Introduction" heading
            body = doc.element.body
            table_elem = target_table._tbl
            intro_para_elem = None

            for para in doc.paragraphs:
                if para.text.strip().startswith('1.') and 'introduction' in para.text.strip().lower():
                    intro_para_elem = para._element
                    break

            if intro_para_elem is not None:
                # Check ordering in body children
                body_children = list(body)
                table_idx = None
                intro_idx = None
                for idx, child in enumerate(body_children):
                    if child is table_elem:
                        table_idx = idx
                    if child is intro_para_elem:
                        intro_idx = idx

                if table_idx is not None and intro_idx is not None and table_idx < intro_idx:
                    print(f"PASS: Component 4 -- Table (idx={table_idx}) appears before Introduction (idx={intro_idx}) (0.20 pts)")
                    total_score += 0.20
                else:
                    print(f"FAIL: Component 4 -- Table position ({table_idx}) not before Introduction ({intro_idx})")
            else:
                print(f"FAIL: Component 4 -- Could not find '1. Introduction' heading to check positioning")
        else:
            print(f"FAIL: Component 4 -- No suitable 4-column table found to check position")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint: persist state then verify
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
