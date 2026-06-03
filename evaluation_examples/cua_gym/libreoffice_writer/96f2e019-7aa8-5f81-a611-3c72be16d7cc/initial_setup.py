"""
Initial Setup: Professional services agreement with six articles (no bookmarks)
Task ID: writer_legal_026
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_026'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Document title --
    title = doc.add_heading('Professional Services Agreement', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        'This Professional Services Agreement ("Agreement") is entered into as of '
        'March 15, 2025, by and between Meridian Consulting Group, LLC ("Provider"), '
        'a limited liability company organized under the laws of the State of Delaware, '
        'with its principal offices at 1200 Market Street, Suite 400, Wilmington, DE 19801, '
        'and Apex Technologies, Inc. ("Client"), a corporation organized under the laws of '
        'the State of California, with its principal offices at 555 Innovation Drive, '
        'San Jose, CA 95134.'
    )
    intro.paragraph_format.space_after = Pt(12)

    preamble = doc.add_paragraph(
        'WHEREAS, the Client desires to engage the Provider to perform certain professional '
        'consulting and advisory services; and WHEREAS, the Provider has the expertise, '
        'personnel, and resources necessary to provide such services; NOW, THEREFORE, in '
        'consideration of the mutual covenants and agreements set forth herein, and for other '
        'good and valuable consideration, the receipt and sufficiency of which are hereby '
        'acknowledged, the parties agree as follows:'
    )
    preamble.paragraph_format.space_after = Pt(16)

    # ---- Article 1: Definitions ----
    doc.add_heading('Article 1 - Definitions', level=1)

    doc.add_paragraph(
        '1.1 "Confidential Information" means any and all non-public, proprietary, or '
        'confidential information disclosed by either party to the other, whether orally, '
        'in writing, or by inspection of tangible objects, including but not limited to '
        'trade secrets, business plans, financial data, customer lists, technical specifications, '
        'and software source code.'
    )
    doc.add_paragraph(
        '1.2 "Deliverables" means all work product, reports, analyses, recommendations, '
        'documentation, software, and other materials created by the Provider in connection '
        'with the Services, as more specifically described in any applicable Statement of Work.'
    )
    doc.add_paragraph(
        '1.3 "Effective Date" means March 15, 2025, the date first written above.'
    )
    doc.add_paragraph(
        '1.4 "Services" means the professional consulting, advisory, and implementation '
        'services to be provided by the Provider to the Client as described in Article 2 '
        'and any applicable Statements of Work executed by the parties.'
    )
    doc.add_paragraph(
        '1.5 "Statement of Work" or "SOW" means a written document executed by both parties '
        'that describes the specific Services to be performed, the timeline, milestones, '
        'Deliverables, and fees associated with a particular engagement.'
    )

    # ---- Article 2: Scope of Services ----
    doc.add_heading('Article 2 - Scope of Services', level=1)

    doc.add_paragraph(
        '2.1 The Provider shall perform the following categories of professional services '
        'for the Client during the term of this Agreement:'
    )
    doc.add_paragraph(
        '(a) Strategic technology assessment and roadmap development, including evaluation '
        'of existing infrastructure, identification of optimization opportunities, and '
        'preparation of a comprehensive modernization plan;'
    )
    doc.add_paragraph(
        '(b) Cloud migration planning and execution support, including architecture design, '
        'risk assessment, data migration strategy, and post-migration validation;'
    )
    doc.add_paragraph(
        '(c) Cybersecurity audit and compliance review, encompassing vulnerability assessment, '
        'penetration testing, policy review, and remediation recommendations aligned with '
        'ISO 27001 and SOC 2 Type II standards;'
    )
    doc.add_paragraph(
        '(d) Staff augmentation and knowledge transfer, providing qualified technical personnel '
        'to supplement the Client\'s team during critical project phases.'
    )
    doc.add_paragraph(
        '2.2 The specific scope, timeline, and deliverables for each engagement shall be '
        'defined in a separate Statement of Work mutually agreed upon by both parties.'
    )

    # ---- Article 3: Payment Terms ----
    doc.add_heading('Article 3 - Payment Terms', level=1)

    doc.add_paragraph(
        '3.1 The Client shall compensate the Provider for Services rendered in accordance '
        'with the fee schedule set forth in the applicable Statement of Work. Unless otherwise '
        'specified, the standard hourly rates are as follows:'
    )
    doc.add_paragraph(
        '(a) Senior Consultant: $275.00 per hour\n'
        '(b) Principal Consultant: $350.00 per hour\n'
        '(c) Managing Director: $450.00 per hour\n'
        '(d) Subject Matter Expert: $500.00 per hour'
    )
    doc.add_paragraph(
        '3.2 The Provider shall submit detailed invoices on a monthly basis, itemizing all '
        'Services performed, hours worked, and expenses incurred. Payment shall be due within '
        'thirty (30) calendar days of the Client\'s receipt of each invoice.'
    )
    doc.add_paragraph(
        '3.3 Late payments shall accrue interest at the rate of 1.5% per month, or the '
        'maximum rate permitted by applicable law, whichever is less, calculated from the '
        'due date until full payment is received.'
    )
    doc.add_paragraph(
        '3.4 Reasonable travel, lodging, and out-of-pocket expenses incurred by the Provider '
        'in connection with the Services shall be reimbursed by the Client, provided such '
        'expenses are pre-approved in writing and supported by appropriate documentation.'
    )

    # ---- Article 4: Termination ----
    doc.add_heading('Article 4 - Termination', level=1)

    doc.add_paragraph(
        '4.1 Either party may terminate this Agreement for convenience upon sixty (60) days\' '
        'prior written notice to the other party.'
    )
    doc.add_paragraph(
        '4.2 Either party may terminate this Agreement immediately upon written notice if '
        'the other party: (a) commits a material breach of any provision of this Agreement '
        'and fails to cure such breach within thirty (30) days after receiving written notice '
        'thereof; (b) becomes insolvent, files for bankruptcy, or has a receiver appointed '
        'for a substantial portion of its assets; or (c) ceases to conduct business in the '
        'normal course.'
    )
    doc.add_paragraph(
        '4.3 Upon termination, the Client shall pay the Provider for all Services rendered '
        'and expenses incurred through the effective date of termination. The Provider shall '
        'deliver to the Client all completed and in-progress Deliverables as of the '
        'termination date.'
    )
    doc.add_paragraph(
        '4.4 The provisions of Articles 1, 5, and 6 shall survive the termination or '
        'expiration of this Agreement.'
    )

    # ---- Article 5: Confidentiality ----
    doc.add_heading('Article 5 - Confidentiality', level=1)

    doc.add_paragraph(
        '5.1 Each party agrees to hold in strict confidence all Confidential Information '
        'received from the other party and to use such information solely for the purpose '
        'of performing its obligations or exercising its rights under this Agreement.'
    )
    doc.add_paragraph(
        '5.2 The receiving party shall protect the disclosing party\'s Confidential Information '
        'using at least the same degree of care it uses to protect its own confidential '
        'information, but in no event less than a reasonable degree of care.'
    )
    doc.add_paragraph(
        '5.3 The obligations of confidentiality shall not apply to information that: '
        '(a) is or becomes publicly available through no fault of the receiving party; '
        '(b) was known to the receiving party prior to disclosure; (c) is independently '
        'developed by the receiving party without use of or reference to the disclosing '
        'party\'s Confidential Information; or (d) is disclosed pursuant to a court order '
        'or governmental requirement, provided the receiving party gives prompt notice to '
        'the disclosing party.'
    )
    doc.add_paragraph(
        '5.4 The confidentiality obligations set forth in this Article shall remain in effect '
        'for a period of five (5) years following the termination or expiration of this Agreement.'
    )

    # ---- Article 6: Miscellaneous ----
    doc.add_heading('Article 6 - Miscellaneous', level=1)

    doc.add_paragraph(
        '6.1 Governing Law. This Agreement shall be governed by and construed in accordance '
        'with the laws of the State of Delaware, without regard to its conflict of law provisions.'
    )
    doc.add_paragraph(
        '6.2 Dispute Resolution. Any dispute arising out of or relating to this Agreement '
        'shall first be submitted to mediation in Wilmington, Delaware. If mediation is '
        'unsuccessful within sixty (60) days, either party may initiate binding arbitration '
        'under the rules of the American Arbitration Association.'
    )
    doc.add_paragraph(
        '6.3 Entire Agreement. This Agreement, together with all Statements of Work and '
        'exhibits, constitutes the entire agreement between the parties and supersedes all '
        'prior negotiations, representations, and agreements relating to the subject matter hereof.'
    )
    doc.add_paragraph(
        '6.4 Amendment. This Agreement may not be modified or amended except by a written '
        'instrument signed by authorized representatives of both parties.'
    )
    doc.add_paragraph(
        '6.5 Notices. All notices required or permitted under this Agreement shall be in '
        'writing and shall be deemed given when delivered personally, sent by certified mail '
        '(return receipt requested), or sent by a nationally recognized overnight courier '
        'service to the addresses set forth in the preamble of this Agreement.'
    )

    # Signature block
    doc.add_paragraph('')
    sig = doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed this Agreement '
                            'as of the date first written above.')
    sig.paragraph_format.space_before = Pt(24)

    doc.add_paragraph('')
    doc.add_paragraph('_______________________________')
    doc.add_paragraph('Meridian Consulting Group, LLC')
    doc.add_paragraph('By: Jonathan R. Whitfield, Managing Partner')
    doc.add_paragraph('')
    doc.add_paragraph('_______________________________')
    doc.add_paragraph('Apex Technologies, Inc.')
    doc.add_paragraph('By: Dr. Priya Sharma, Chief Technology Officer')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
