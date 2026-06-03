"""
Reward Script: Weekly Meal Plan Template
Task ID: writer_creative_062
Domain: libreoffice_writer
Scoring:
  Component 1: Title 'Weekly Meal Plan' — font_size=20pt, bold, centered (0.25 pts)
  Component 2: Table structure — 5x8, correct header (Meal/Mon-Sun) + meal rows, #AED581 bg,
               header row and first column bold (0.30 pts)
  Component 3: 'Grocery List' heading — 16pt, bold, space_before=18pt, 10 bulleted blank lines (0.25 pts)
  Component 4: 'Week of' line — centered, space_after=12pt (0.20 pts)
  Total: 1.0
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_062'

def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title 'Weekly Meal Plan' — font_size=20pt, bold=True, alignment=CENTER (0.25 pts)
    # This fails on initial (12pt, not bold, not centered) and passes on golden (20pt, bold, center).
    try:
        title_para = None
        for para in doc.paragraphs:
            if 'Weekly Meal Plan' in para.text:
                title_para = para
                break

        if title_para is None:
            print("FAIL: Component 1 — 'Weekly Meal Plan' paragraph not found")
        else:
            # Check alignment
            alignment_ok = title_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER

            # Check runs for bold and size
            title_bold = False
            title_size_ok = False
            for run in title_para.runs:
                if run.text.strip():
                    if run.bold is True:
                        title_bold = True
                    if run.font.size and abs(run.font.size.pt - 20.0) < 0.5:
                        title_size_ok = True

            if alignment_ok and title_bold and title_size_ok:
                print(f"PASS: Component 1 — Title is 20pt, bold, centered (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 — Title formatting: aligned={alignment_ok}, bold={title_bold}, size_20pt={title_size_ok}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table structure — 5 rows x 8 cols, correct header, #AED581 bg, bold first col (0.30 pts)
    # This fails on initial (no table) and passes on golden.
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 2 — No table found in document")
        else:
            table = doc.tables[0]
            rows = len(table.rows)
            cols = len(table.columns)

            # Verify dimensions: 5 rows (header + 4 meal types) x 8 cols (Meal + Mon-Sun)
            dims_ok = (rows == 5 and cols == 8)

            # Verify header row text
            expected_header = ['Meal', 'Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
            header_texts = [table.cell(0, j).text.strip() for j in range(cols)] if dims_ok else []
            header_text_ok = (header_texts == expected_header)

            # Verify header row background color = AED581
            header_bg_ok = True
            if dims_ok:
                for j in range(cols):
                    tc = table.cell(0, j)._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            if fill and fill.upper() != 'AED581':
                                header_bg_ok = False
                                break
                        else:
                            header_bg_ok = False
                            break
                    else:
                        header_bg_ok = False
                        break

            # Verify header row bold
            header_bold_ok = True
            if dims_ok:
                for j in range(cols):
                    cell = table.cell(0, j)
                    cell_bold = False
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip() and run.bold:
                                cell_bold = True
                    if not cell_bold:
                        header_bold_ok = False
                        break

            # Verify first column meal labels and bold
            expected_meals = ['Breakfast', 'Lunch', 'Dinner', 'Snacks']
            meal_col_ok = True
            if dims_ok:
                for i, meal in enumerate(expected_meals, start=1):
                    cell_text = table.cell(i, 0).text.strip()
                    if cell_text != meal:
                        meal_col_ok = False
                        break
                    # Check bold
                    cell_bold = False
                    for para in table.cell(i, 0).paragraphs:
                        for run in para.runs:
                            if run.text.strip() and run.bold:
                                cell_bold = True
                    if not cell_bold:
                        meal_col_ok = False
                        break

            if dims_ok and header_text_ok and header_bg_ok and header_bold_ok and meal_col_ok:
                print(f"PASS: Component 2 — Table 5x8, correct header, #AED581 bg, bold first col (0.30 pts)")
                total_score += 0.30
            else:
                print(f"FAIL: Component 2 — dims_ok={dims_ok} ({rows}x{cols}), header_text_ok={header_text_ok}, "
                      f"header_bg_ok={header_bg_ok}, header_bold_ok={header_bold_ok}, meal_col_ok={meal_col_ok}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: 'Grocery List' heading — 16pt, bold, space_before=18pt, 10 bulleted blank lines (0.25 pts)
    # This fails on initial (12pt, not bold, no bullets) and passes on golden.
    try:
        grocery_para = None
        grocery_idx = None
        for i, para in enumerate(doc.paragraphs):
            if 'Grocery List' in para.text:
                grocery_para = para
                grocery_idx = i
                break

        if grocery_para is None:
            print("FAIL: Component 3 — 'Grocery List' paragraph not found")
        else:
            # Check size and bold
            grocery_bold = False
            grocery_size_ok = False
            for run in grocery_para.runs:
                if run.text.strip():
                    if run.bold is True:
                        grocery_bold = True
                    if run.font.size and abs(run.font.size.pt - 16.0) < 0.5:
                        grocery_size_ok = True

            # Check space_before (~18pt = 228600 EMU, allow tolerance)
            sb = grocery_para.paragraph_format.space_before
            space_before_ok = (sb is not None and abs(sb - Pt(18)) < Pt(2))

            # Check 10 bullet lines below
            bullet_count = 0
            if grocery_idx is not None:
                for para in doc.paragraphs[grocery_idx + 1:]:
                    if para.style.name in ('List Bullet', 'List Bullet 2', 'List Bullet 3') or \
                       'List Bullet' in para.style.name:
                        bullet_count += 1

            bullets_ok = (bullet_count >= 10)

            if grocery_bold and grocery_size_ok and space_before_ok and bullets_ok:
                print(f"PASS: Component 3 — Grocery List 16pt bold, space_before=18pt, {bullet_count} bullets (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — bold={grocery_bold}, size_16pt={grocery_size_ok}, "
                      f"space_before_ok={space_before_ok} (sb={sb}), bullets={bullet_count}>=10:{bullets_ok}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: 'Week of: ______________' — centered, space_after=12pt (0.20 pts)
    # This fails on initial (not centered, no space_after) and passes on golden.
    try:
        weekof_para = None
        for para in doc.paragraphs:
            if 'Week of:' in para.text:
                weekof_para = para
                break

        if weekof_para is None:
            print("FAIL: Component 4 — 'Week of:' paragraph not found")
        else:
            alignment_ok = weekof_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
            sa = weekof_para.paragraph_format.space_after
            space_after_ok = (sa is not None and abs(sa - Pt(12)) < Pt(2))

            if alignment_ok and space_after_ok:
                print(f"PASS: Component 4 — 'Week of' line centered with space_after=12pt (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 4 — alignment_ok={alignment_ok}, space_after_ok={space_after_ok} (sa={sa})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/meal_plan_template.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
