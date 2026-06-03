"""
Initial Setup: Weekly Meal Plan Template - Plain Text Initial State
Task ID: writer_creative_062
Domain: libreoffice_writer

Creates a plain text document with meal planning content — no formatting,
no table, no bullets. The agent must apply all formatting and structure.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_062'
OUTPUT = f'{WORKDIR}/meal_plan_template.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title — plain text, 12pt, left-aligned (no bold, no centering)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('Weekly Meal Plan')
    title_run.font.size = Pt(12)
    title_run.bold = False
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Week of line — plain text, 12pt, left-aligned
    week_para = doc.add_paragraph()
    week_run = week_para.add_run('Week of: ______________')
    week_run.font.size = Pt(12)
    week_run.bold = False
    week_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Meal types listed as plain text (no table)
    meals_para = doc.add_paragraph()
    meals_run = meals_para.add_run('Breakfast, Lunch, Dinner, Snacks')
    meals_run.font.size = Pt(12)
    meals_run.bold = False
    meals_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Days of week as plain text (no table)
    days_para = doc.add_paragraph()
    days_run = days_para.add_run('Monday, Tuesday, Wednesday, Thursday, Friday, Saturday, Sunday')
    days_run.font.size = Pt(12)
    days_run.bold = False
    days_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Grocery List heading — plain text, 12pt, left-aligned (no bold, no font sizing)
    grocery_para = doc.add_paragraph()
    grocery_run = grocery_para.add_run('Grocery List')
    grocery_run.font.size = Pt(12)
    grocery_run.bold = False
    grocery_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
