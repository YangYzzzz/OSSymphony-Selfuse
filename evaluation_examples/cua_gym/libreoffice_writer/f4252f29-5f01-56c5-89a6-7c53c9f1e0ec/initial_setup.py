"""
Initial Setup: Merge three reviewer comment documents into thesis chapter
Task ID: writer_acad_094
Domain: libreoffice_writer

Creates:
  - /home/user/writer_acad_094.docx  (main thesis chapter, clean, no tracked changes)
  - /home/user/reviewer1.docx  (reviewer 1's tracked changes & comments)
  - /home/user/reviewer2.docx  (reviewer 2's tracked changes & comments)
  - /home/user/reviewer3.docx  (reviewer 3's tracked changes & comments)

Opens writer_acad_094.docx in LibreOffice Writer.
"""

import os
import shlex
import subprocess
import time
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree
import datetime

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_094'
MAIN_DOC = f'{WORKDIR}/{TASK_ID}.docx'
REV1_DOC = f'{WORKDIR}/reviewer1.docx'
REV2_DOC = f'{WORKDIR}/reviewer2.docx'
REV3_DOC = f'{WORKDIR}/reviewer3.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# ---------- Thesis chapter content ----------

CHAPTER_TITLE = "Chapter 3: Computational Approaches to Protein Folding Prediction"

SECTIONS = [
    {
        "heading": "3.1 Introduction",
        "paragraphs": [
            "Protein folding remains one of the most challenging problems in computational biology. "
            "The ability to predict three-dimensional protein structures from amino acid sequences has "
            "profound implications for drug discovery, enzyme engineering, and our fundamental understanding "
            "of molecular biology. Since Anfinsen's seminal experiments in the 1970s demonstrating that the "
            "amino acid sequence contains sufficient information to determine a protein's native conformation, "
            "researchers have pursued increasingly sophisticated computational methods to solve this problem.",

            "This chapter examines the evolution of computational approaches to protein folding prediction, "
            "from early energy minimization methods through modern deep learning architectures. We evaluate "
            "the strengths and limitations of each approach and propose a hybrid framework that combines "
            "physics-based modeling with data-driven techniques to achieve improved prediction accuracy "
            "across diverse protein families.",

            "The remainder of this chapter is organized as follows. Section 3.2 reviews classical energy-based "
            "methods. Section 3.3 discusses template-based and homology modeling. Section 3.4 examines modern "
            "machine learning approaches including AlphaFold and ESMFold. Section 3.5 presents our proposed "
            "hybrid framework, and Section 3.6 discusses experimental validation results."
        ]
    },
    {
        "heading": "3.2 Classical Energy-Based Methods",
        "paragraphs": [
            "The earliest computational approaches to protein structure prediction relied on molecular dynamics "
            "simulations and energy minimization algorithms. These methods attempt to find the global minimum "
            "of a protein's free energy landscape, guided by physical force fields such as AMBER, CHARMM, and "
            "OPLS-AA. The fundamental assumption is that a protein's native state corresponds to the thermodynamic "
            "minimum of the system.",

            "Molecular dynamics simulations numerically integrate Newton's equations of motion for all atoms in "
            "the system, including explicit solvent molecules and counterions. While this approach captures the "
            "physics of folding with high fidelity, the computational cost is prohibitive for most proteins of "
            "biological interest. The folding timescale for typical proteins ranges from microseconds to seconds, "
            "while current hardware limits all-atom simulations to microsecond timescales for systems containing "
            "fewer than 100,000 atoms.",

            "Several coarse-grained approaches have been developed to address this limitation. The MARTINI force "
            "field maps approximately four heavy atoms to a single interaction site, reducing computational cost "
            "by two to three orders of magnitude. Similarly, Go-model approaches use native contacts as the sole "
            "basis for interaction potentials, enabling rapid sampling of folding landscapes at the cost of "
            "reduced transferability to non-native conformations.",

            "Despite significant advances in hardware acceleration, including the Anton supercomputer designed "
            "specifically for molecular dynamics, energy-based methods remain impractical for proteome-scale "
            "structure prediction. Their primary value lies in studying folding mechanisms and refining structures "
            "obtained through other methods."
        ]
    },
    {
        "heading": "3.3 Template-Based and Homology Modeling",
        "paragraphs": [
            "Template-based modeling exploits the observation that evolutionarily related proteins adopt similar "
            "three-dimensional structures. Homology modeling, the most widely used structural biology technique, "
            "builds a target protein model using experimentally determined structures of homologous proteins as "
            "templates. The critical steps include template identification, sequence-structure alignment, backbone "
            "generation, loop modeling, and side-chain placement.",

            "The accuracy of homology models depends primarily on the sequence identity between target and template. "
            "Models built from templates sharing greater than 50% sequence identity typically achieve root-mean-square "
            "deviation (RMSD) values below 1.5 angstroms relative to experimental structures. Below 30% sequence "
            "identity, alignment errors become significant and model quality degrades substantially.",

            "Threading methods, also known as fold recognition, extend template-based approaches to cases of remote "
            "homology where sequence similarity is undetectable. These methods evaluate the compatibility of a target "
            "sequence with known protein folds using statistical potentials derived from the Protein Data Bank. "
            "Programs such as I-TASSER, Phyre2, and HHpred have demonstrated remarkable success in CASP experiments, "
            "particularly for targets with detectable structural homologs."
        ]
    },
    {
        "heading": "3.4 Machine Learning Approaches",
        "paragraphs": [
            "The application of machine learning to protein structure prediction has transformed the field. Early "
            "approaches used neural networks to predict secondary structure elements and contact maps from sequence "
            "features. The breakthrough came with the development of attention-based architectures that directly "
            "predict three-dimensional coordinates from multiple sequence alignments.",

            "AlphaFold2, developed by DeepMind, achieved unprecedented accuracy in the CASP14 competition, with "
            "median GDT-TS scores exceeding 90 for most targets. The architecture combines an evoformer module that "
            "processes multiple sequence alignment information with a structure module that iteratively refines "
            "atomic coordinates. The model's key innovation is the use of invariant point attention, which operates "
            "directly on three-dimensional coordinates while respecting SE(3) equivariance.",

            "ESMFold and other language model approaches demonstrate that accurate structure prediction is possible "
            "without explicit multiple sequence alignments. These models learn evolutionary information implicitly "
            "through pretraining on large sequence databases. While single-sequence methods currently lag behind "
            "MSA-based approaches for most targets, they offer significant speed advantages and perform well on "
            "orphan proteins lacking homologous sequences.",

            "Despite their remarkable success, current deep learning methods have notable limitations. Prediction "
            "accuracy degrades for intrinsically disordered regions, multimeric complexes, and proteins undergoing "
            "large conformational changes. Additionally, confidence estimates, while generally well-calibrated, "
            "can be misleading for certain classes of membrane proteins and repeat proteins."
        ]
    },
    {
        "heading": "3.5 Proposed Hybrid Framework",
        "paragraphs": [
            "We propose a hybrid framework that leverages the complementary strengths of physics-based and "
            "data-driven methods. Our approach, termed PhysML-Fold, consists of three stages: initial structure "
            "generation using an enhanced transformer architecture, physics-based refinement using adaptive "
            "molecular dynamics, and ensemble selection guided by free energy perturbation calculations.",

            "In the first stage, we modify the AlphaFold2 architecture to incorporate explicit physical constraints "
            "during training. Bond lengths, bond angles, and Ramachandran plot distributions are enforced through "
            "additional loss terms weighted by a curriculum schedule. This produces initial models with improved "
            "stereochemistry compared to the baseline architecture.",

            "The refinement stage employs replica exchange molecular dynamics with an adaptive biasing potential "
            "derived from the neural network's confidence estimates. Regions with low predicted local distance "
            "difference test (pLDDT) scores receive enhanced sampling, while high-confidence regions are restrained "
            "to preserve favorable contacts. This targeted approach reduces computational cost by focusing simulation "
            "effort where it is most needed.",

            "Final model selection uses free energy perturbation calculations to rank ensemble members. We compute "
            "relative binding free energies for a panel of small-molecule probes to assess the quality of predicted "
            "binding sites, providing an orthogonal quality metric to geometry-based scores."
        ]
    },
    {
        "heading": "3.6 Experimental Validation",
        "paragraphs": [
            "We evaluated PhysML-Fold on a benchmark set of 150 protein domains from the CASP15 competition. "
            "Table 3.1 summarizes the results across different target categories. Our method achieved a median "
            "GDT-TS of 82.4 across all targets, compared to 79.1 for standard AlphaFold2 and 71.3 for Rosetta "
            "ab initio predictions.",

            "The improvement was most pronounced for targets in the difficult free-modeling category, where "
            "PhysML-Fold achieved a median GDT-TS of 64.7 compared to 58.2 for AlphaFold2. This suggests that "
            "physics-based refinement adds the most value when initial predictions are uncertain. For easy "
            "template-based targets, the two methods performed comparably.",

            "Binding site quality, assessed using small-molecule docking enrichment factors, showed a 23% "
            "improvement over unrefined AlphaFold2 predictions. This is particularly relevant for drug discovery "
            "applications, where accurate active site geometry is essential for virtual screening campaigns.",

            "We note several limitations of our evaluation. The benchmark set, while diverse, may not fully "
            "represent the distribution of prediction targets encountered in practice. Additionally, the "
            "computational cost of the refinement stage (approximately 4 GPU-hours per target) may limit "
            "applicability for large-scale proteome annotation."
        ]
    },
    {
        "heading": "3.7 Conclusions",
        "paragraphs": [
            "This chapter presented a comprehensive review of computational protein structure prediction methods "
            "and introduced PhysML-Fold, a hybrid framework combining deep learning with physics-based refinement. "
            "Our results demonstrate that integrating physical constraints into neural network training and "
            "applying targeted molecular dynamics refinement can improve prediction accuracy, particularly for "
            "challenging targets.",

            "Future work will focus on extending PhysML-Fold to multimeric complex prediction and incorporating "
            "experimental restraints from cross-linking mass spectrometry and cryo-electron microscopy data. "
            "We anticipate that the convergence of machine learning and physics-based methods will continue to "
            "advance our ability to predict and understand protein structures at atomic resolution."
        ]
    }
]


def create_main_document():
    """Create the clean thesis chapter document."""
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Title
    title = doc.add_heading(CHAPTER_TITLE, level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Author line
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author_para.add_run("Elena Vasquez")
    run.font.size = Pt(11)
    run.italic = True

    # Affiliation
    affil_para = doc.add_paragraph()
    affil_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = affil_para.add_run("Department of Computational Biology, Stanford University")
    run.font.size = Pt(10)

    # Abstract
    doc.add_heading("Abstract", level=2)
    abstract_para = doc.add_paragraph(
        "We present PhysML-Fold, a hybrid computational framework for protein structure prediction "
        "that combines deep learning architectures with physics-based molecular dynamics refinement. "
        "Our approach achieves a median GDT-TS of 82.4 on CASP15 benchmark targets, representing a "
        "significant improvement over existing methods, particularly for difficult free-modeling targets. "
        "We demonstrate that integrating physical constraints during neural network training and applying "
        "targeted refinement to low-confidence regions yields more accurate structures with improved "
        "binding site geometry."
    )
    abstract_para.paragraph_format.first_line_indent = Inches(0.5)

    # Sections
    for section in SECTIONS:
        doc.add_heading(section["heading"], level=2)
        for para_text in section["paragraphs"]:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.first_line_indent = Inches(0.5)

    # References section
    doc.add_heading("References", level=2)
    refs = [
        "Anfinsen, C.B. (1973). Principles that govern the folding of protein chains. Science, 181(4096), 223-230.",
        "Jumper, J. et al. (2021). Highly accurate protein structure prediction with AlphaFold. Nature, 596, 583-589.",
        "Lin, Z. et al. (2023). Evolutionary-scale prediction of atomic-level protein structure with a language model. Science, 379(6637), 1123-1130.",
        "Baek, M. et al. (2021). Accurate prediction of protein structures and interactions using a three-track neural network. Science, 373(6557), 871-876.",
        "Shaw, D.E. et al. (2010). Atomic-level characterization of the structural dynamics of proteins. Science, 330(6002), 341-346.",
        "Kryshtafovych, A. et al. (2023). Critical assessment of methods of protein structure prediction (CASP15). Proteins, 91(12), 1615-1624.",
        "Souza, P.C.T. et al. (2021). Martini 3: a general purpose force field for coarse-grained molecular dynamics. Nature Methods, 18(4), 382-388.",
        "Yang, J. et al. (2020). Improved protein structure prediction using predicted interresidue orientations. PNAS, 117(3), 1496-1503."
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"[{i}] {ref}")
        p.paragraph_format.space_after = Pt(2)

    doc.save(MAIN_DOC)
    print(f'Main thesis document created: {MAIN_DOC}')


def add_comment(doc, paragraph, comment_text, author, date_str, initials, comment_id):
    """Add a comment to a paragraph using OOXML manipulation."""
    # Add to comments part if not present
    # We need to work at the XML level for comments

    # Create comment range start
    comment_start = OxmlElement('w:commentRangeStart')
    comment_start.set(qn('w:id'), str(comment_id))

    # Insert at beginning of paragraph
    paragraph._element.insert(0, comment_start)

    # Create comment range end
    comment_end = OxmlElement('w:commentRangeEnd')
    comment_end.set(qn('w:id'), str(comment_id))
    paragraph._element.append(comment_end)

    # Create comment reference run
    comment_ref_run = OxmlElement('w:r')
    comment_ref_rpr = OxmlElement('w:rPr')
    comment_ref_style = OxmlElement('w:rStyle')
    comment_ref_style.set(qn('w:val'), 'CommentReference')
    comment_ref_rpr.append(comment_ref_style)
    comment_ref_run.append(comment_ref_rpr)
    comment_ref_mark = OxmlElement('w:commentReference')
    comment_ref_mark.set(qn('w:id'), str(comment_id))
    comment_ref_run.append(comment_ref_mark)
    paragraph._element.append(comment_ref_run)


def add_tracked_insertion(paragraph, text, author, date_str, rev_id):
    """Add a tracked insertion (new text) to a paragraph."""
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(rev_id))
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), date_str)

    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    run.append(rpr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    ins.append(run)

    paragraph._element.append(ins)


def add_tracked_deletion(paragraph, text, author, date_str, rev_id):
    """Add a tracked deletion (removed text) to a paragraph."""
    delete = OxmlElement('w:del')
    delete.set(qn('w:id'), str(rev_id))
    delete.set(qn('w:author'), author)
    delete.set(qn('w:date'), date_str)

    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    run.append(rpr)
    del_text = OxmlElement('w:delText')
    del_text.set(qn('xml:space'), 'preserve')
    del_text.text = text
    run.append(del_text)
    delete.append(run)

    paragraph._element.append(delete)


def create_comments_part(doc, comments_data):
    """Create the comments XML part in the document."""
    from docx.opc.part import Part
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    # Build comments XML
    WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    comments_xml = f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    comments_xml += f'<w:comments xmlns:w="{WORD_NS}">'

    for c in comments_data:
        comments_xml += (
            f'<w:comment w:id="{c["id"]}" w:author="{c["author"]}" '
            f'w:date="{c["date"]}" w:initials="{c["initials"]}">'
            f'<w:p><w:r><w:t>{c["text"]}</w:t></w:r></w:p>'
            f'</w:comment>'
        )

    comments_xml += '</w:comments>'

    # Add as part
    partname = '/word/comments.xml'
    content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'

    from docx.opc.part import Part as OpcPart
    from docx.opc.packuri import PackURI

    comments_part = OpcPart(
        PackURI(partname),
        content_type,
        comments_xml.encode('utf-8'),
        doc.part.package,
    )

    # Add relationship
    COMMENTS_REL_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
    doc.part.relate_to(comments_part, COMMENTS_REL_TYPE)


def create_reviewer_doc(base_doc_path, output_path, reviewer_name, reviewer_initials,
                        date_str, changes):
    """
    Create a reviewer document based on the thesis chapter with tracked changes and comments.

    changes is a list of dicts:
      - type: 'insert', 'delete', or 'comment'
      - para_index: which paragraph (by index in doc.paragraphs)
      - text: text to insert/delete or comment content
      - position: 'end' or 'replace' (for inserts)
    """
    import shutil
    shutil.copy(base_doc_path, output_path)
    doc = Document(output_path)

    comments_data = []
    comment_id_counter = 100
    rev_id_counter = 200

    for change in changes:
        para_idx = change['para_index']
        if para_idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[para_idx]

        if change['type'] == 'insert':
            add_tracked_insertion(para, change['text'], reviewer_name, date_str, rev_id_counter)
            rev_id_counter += 1

        elif change['type'] == 'delete':
            add_tracked_deletion(para, change['text'], reviewer_name, date_str, rev_id_counter)
            rev_id_counter += 1

        elif change['type'] == 'comment':
            add_comment(doc, para, change['text'], reviewer_name, date_str,
                       reviewer_initials, comment_id_counter)
            comments_data.append({
                'id': comment_id_counter,
                'author': reviewer_name,
                'date': date_str,
                'initials': reviewer_initials,
                'text': change['text'],
            })
            comment_id_counter += 1

    # Add comments part if we have comments
    if comments_data:
        create_comments_part(doc, comments_data)

    doc.save(output_path)
    print(f'Reviewer document created: {output_path}')


def create_reviewer_documents():
    """Create all three reviewer documents with tracked changes and comments."""

    # Reviewer 1: Dr. James Thornton (methodology expert)
    # Focuses on Section 3.2 (Classical Methods) and 3.5 (Proposed Framework)
    reviewer1_changes = [
        {
            'type': 'comment',
            'para_index': 5,  # First para of Section 3.2
            'text': 'Consider citing the GROMOS force field as well, which is widely used in European research groups.',
        },
        {
            'type': 'insert',
            'para_index': 6,  # Second para of 3.2
            'text': ' Recent work by Lindorff-Larsen et al. (2024) has demonstrated improved sampling efficiency using machine-learned force fields.',
        },
        {
            'type': 'comment',
            'para_index': 7,  # Third para of 3.2
            'text': 'The MARTINI force field description should mention version 3 specifically, as it represents a significant improvement over earlier versions.',
        },
        {
            'type': 'delete',
            'para_index': 8,  # Fourth para of 3.2
            'text': 'impractical',
        },
        {
            'type': 'insert',
            'para_index': 8,
            'text': 'computationally demanding',
        },
        {
            'type': 'comment',
            'para_index': 18,  # In Section 3.5
            'text': 'The curriculum schedule for physics constraints needs more detail. What are the specific weighting factors and at what training epoch do you transition?',
        },
        {
            'type': 'insert',
            'para_index': 19,  # Refinement stage in 3.5
            'text': ' We recommend including a comparison with standard Rosetta relaxation protocols to contextualize the computational overhead.',
        },
    ]

    # Reviewer 2: Prof. Mei-Ling Chen (ML/AI expert)
    # Focuses on Section 3.4 (ML Approaches) and validation
    reviewer2_changes = [
        {
            'type': 'comment',
            'para_index': 13,  # First para of 3.4
            'text': 'This section should discuss the role of multiple sequence alignment depth on prediction quality. Recent studies show diminishing returns beyond 1000 effective sequences.',
        },
        {
            'type': 'insert',
            'para_index': 14,  # AlphaFold2 para
            'text': ' It is worth noting that AlphaFold3 has since extended this architecture to handle nucleic acids and small molecules within a unified framework.',
        },
        {
            'type': 'delete',
            'para_index': 15,  # ESMFold para
            'text': 'currently lag behind',
        },
        {
            'type': 'insert',
            'para_index': 15,
            'text': 'show competitive but slightly lower performance compared to',
        },
        {
            'type': 'comment',
            'para_index': 16,  # Limitations para of 3.4
            'text': 'Consider adding a discussion of hallucination in structure prediction - cases where high-confidence predictions are demonstrably wrong.',
        },
        {
            'type': 'comment',
            'para_index': 22,  # Validation results
            'text': 'The benchmark set of 150 domains may be insufficient for statistical significance. Consider bootstrapping confidence intervals for the GDT-TS comparisons.',
        },
        {
            'type': 'insert',
            'para_index': 23,  # Binding site quality
            'text': ' Furthermore, the docking enrichment factors should be reported with standard deviations across the target set.',
        },
    ]

    # Reviewer 3: Dr. Aisha Patel (structural biology expert)
    # Focuses on experimental validation and biological context
    reviewer3_changes = [
        {
            'type': 'comment',
            'para_index': 3,  # End of introduction
            'text': 'The introduction should mention the impact of cryo-EM resolution revolution on the training data available for ML methods.',
        },
        {
            'type': 'insert',
            'para_index': 9,  # Template-based modeling first para
            'text': ' It is important to note that the quality of template structures in the PDB has improved significantly with advances in cryo-electron microscopy.',
        },
        {
            'type': 'comment',
            'para_index': 10,  # Accuracy of homology models
            'text': 'The 50% sequence identity threshold is somewhat dated. Recent comparative studies suggest 40% may be more appropriate with modern alignment algorithms.',
        },
        {
            'type': 'delete',
            'para_index': 10,
            'text': '50%',
        },
        {
            'type': 'insert',
            'para_index': 10,
            'text': '40-50%',
        },
        {
            'type': 'comment',
            'para_index': 22,  # Validation first para
            'text': 'Were any experimental validation experiments performed (e.g., SAXS, crosslinking-MS)? Computational benchmarks alone may not be convincing to structural biologists.',
        },
        {
            'type': 'insert',
            'para_index': 24,  # Limitations
            'text': ' We also acknowledge that the refinement protocol has not been tested on membrane proteins, which represent approximately 30% of the human proteome.',
        },
        {
            'type': 'comment',
            'para_index': 25,  # Conclusions first para
            'text': 'The conclusions should more explicitly state the practical implications for structural biology workflows and drug discovery timelines.',
        },
    ]

    create_reviewer_doc(
        MAIN_DOC, REV1_DOC,
        "Dr. James Thornton", "JT",
        "2025-11-15T10:30:00Z",
        reviewer1_changes
    )

    create_reviewer_doc(
        MAIN_DOC, REV2_DOC,
        "Prof. Mei-Ling Chen", "MC",
        "2025-11-18T14:22:00Z",
        reviewer2_changes
    )

    create_reviewer_doc(
        MAIN_DOC, REV3_DOC,
        "Dr. Aisha Patel", "AP",
        "2025-11-20T09:15:00Z",
        reviewer3_changes
    )


def main():
    create_main_document()
    create_reviewer_documents()

    # Open main document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{MAIN_DOC}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
