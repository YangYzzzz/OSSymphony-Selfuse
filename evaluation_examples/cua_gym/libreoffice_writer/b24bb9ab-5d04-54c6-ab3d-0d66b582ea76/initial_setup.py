"""
Initial Setup: Create a business report with uppercase words to be converted to title case.
Task ID: writer_af_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_af_043'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Title Page ---
    doc.add_heading('Quarterly PERFORMANCE Review Report', level=0)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('DEPARTMENT of Strategic OPERATIONS')
    run.font.size = Pt(16)
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p2.add_run('Fiscal Year 2025 - Q1 ASSESSMENT')
    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p3.add_run('Prepared by the MANAGEMENT Team')
    doc.add_page_break()

    # --- Page 2: Executive Summary ---
    doc.add_heading('Executive SUMMARY', level=1)
    doc.add_paragraph(
        'This ANALYSIS provides a comprehensive overview of the organizational '
        'PERFORMANCE metrics collected during the first quarter of fiscal year 2025. '
        'The COMMITTEE has reviewed all submissions from each DEPARTMENT and compiled '
        'the findings into this consolidated report for senior LEADERSHIP review.'
    )
    doc.add_paragraph(
        'Key areas of focus include REVENUE growth, EMPLOYEE satisfaction, '
        'operational EFFICIENCY, and market EXPANSION strategies. The IT and AI '
        'teams contributed supporting data for this assessment.'
    )
    doc.add_page_break()

    # --- Page 3: Financial Overview ---
    doc.add_heading('Financial OVERVIEW', level=1)
    doc.add_paragraph(
        'The FINANCE division reported strong REVENUE numbers across all major '
        'product lines. Total REVENUE for Q1 reached $12.4 million, representing '
        'a 15% increase over the previous QUARTER. The BUDGET allocated for '
        'MARKETING activities was fully utilized, generating substantial returns.'
    )
    doc.add_heading('Revenue Breakdown by SEGMENT', level=2)

    # Table for financial data
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['SEGMENT', 'Q1 Revenue', 'Q4 Previous', 'Growth %']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    data = [
        ['Enterprise Solutions', '$4,230,000', '$3,780,000', '11.9%'],
        ['Consumer Products', '$3,150,000', '$2,890,000', '9.0%'],
        ['CONSULTING Services', '$2,870,000', '$2,410,000', '19.1%'],
        ['Cloud INFRASTRUCTURE', '$1,450,000', '$1,120,000', '29.5%'],
        ['Professional TRAINING', '$700,000', '$580,000', '20.7%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_page_break()

    # --- Page 4: Operations ---
    doc.add_heading('OPERATIONS and LOGISTICS', level=1)
    doc.add_paragraph(
        'The OPERATIONS team achieved significant milestones during Q1. '
        'Supply chain OPTIMIZATION efforts led to a 22% reduction in lead times. '
        'The WAREHOUSE management system upgrade was completed on schedule, '
        'and the new INVENTORY tracking protocol is now fully operational.'
    )
    doc.add_paragraph(
        'LOGISTICS improvements included the deployment of automated routing '
        'ALGORITHMS that reduced delivery costs by 18%. The TRANSPORTATION '
        'fleet was expanded with 12 new vehicles to support growing demand.'
    )
    doc.add_page_break()

    # --- Page 5: Human Resources ---
    doc.add_heading('Human RESOURCES Update', level=1)
    doc.add_paragraph(
        'The RECRUITMENT drive in Q1 successfully onboarded 47 new employees '
        'across all departments. The COMPLIANCE training program achieved 98% '
        'completion rate. Employee SATISFACTION scores improved by 8 points '
        'compared to the previous survey period.'
    )
    doc.add_paragraph(
        'The COMPENSATION review was conducted in February, with adjustments '
        'effective from March. The BENEFITS package was enhanced to include '
        'additional wellness programs. RETENTION rates remain above industry '
        'average at 94.2%.'
    )
    doc.add_page_break()

    # --- Page 6: Technology ---
    doc.add_heading('TECHNOLOGY and INNOVATION', level=1)
    doc.add_paragraph(
        'The ENGINEERING team delivered three major platform updates during Q1. '
        'CYBERSECURITY measures were strengthened following the annual audit. '
        'The ARCHITECTURE review recommended migration to a microservices '
        'model, which is now in DEVELOPMENT.'
    )
    doc.add_paragraph(
        'The AI and IT departments collaborated on a machine learning IMPLEMENTATION '
        'for customer service automation. INTEGRATION with existing CRM systems '
        'is expected to complete by the end of Q2.'
    )
    doc.add_page_break()

    # --- Page 7: Marketing ---
    doc.add_heading('MARKETING and Communications', level=1)
    doc.add_paragraph(
        'The ADVERTISING campaign launched in January exceeded all EXPECTATIONS. '
        'Brand AWARENESS metrics showed a 34% improvement in target demographics. '
        'The STRATEGY team developed a comprehensive social media plan that '
        'increased ENGAGEMENT by 45% across all platforms.'
    )
    doc.add_paragraph(
        'COMMUNICATIONS efforts focused on internal NEWSLETTER improvements '
        'and external press COVERAGE. The BRANDING refresh project is on track '
        'for completion in Q2.'
    )
    doc.add_page_break()

    # --- Page 8: Risk and Compliance ---
    doc.add_heading('Risk MANAGEMENT and COMPLIANCE', level=1)
    doc.add_paragraph(
        'The GOVERNANCE framework was updated to reflect new regulatory requirements. '
        'All DOCUMENTATION was reviewed and approved by the legal team. '
        'The CERTIFICATION process for ISO 27001 is progressing as planned, '
        'with the external ASSESSMENT scheduled for June.'
    )
    doc.add_paragraph(
        'REGULATORY changes in the financial sector require additional '
        'MONITORING of cross-border transactions. The SURVEILLANCE system '
        'upgrade was approved and is in the PROCUREMENT phase.'
    )
    doc.add_page_break()

    # --- Page 9: Strategic Outlook ---
    doc.add_heading('Strategic OUTLOOK', level=1)
    doc.add_paragraph(
        'The EXECUTIVE team has outlined three priority areas for Q2: '
        'EXPANSION into the Southeast Asian market, CONSOLIDATION of '
        'existing product lines, and ACCELERATION of digital transformation '
        'initiatives. The INVESTMENT committee approved $2.8 million for '
        'these strategic priorities.'
    )
    doc.add_paragraph(
        'PARTNERSHIP opportunities with two major DISTRIBUTION networks '
        'are under NEGOTIATION. The ACQUISITION target list has been '
        'narrowed to three candidates for due DILIGENCE review.'
    )
    doc.add_page_break()

    # --- Page 10: Conclusion ---
    doc.add_heading('CONCLUSION and RECOMMENDATIONS', level=1)
    doc.add_paragraph(
        'In CONCLUSION, the first quarter has demonstrated strong MOMENTUM '
        'across all business units. The ORGANIZATION is well-positioned to '
        'achieve its annual OBJECTIVES. The following RECOMMENDATIONS are '
        'submitted for CONSIDERATION by senior leadership:'
    )

    recs = [
        'Increase ALLOCATION for cloud infrastructure by 20%',
        'Expand the RECRUITMENT pipeline for engineering roles',
        'Accelerate the MIGRATION to the new ERP system',
        'Enhance COLLABORATION tools for remote team members',
    ]
    for rec in recs:
        doc.add_paragraph(rec, style='List Bullet')

    doc.add_paragraph(
        'This report was prepared by the MANAGEMENT COMMITTEE and reviewed '
        'by the GOVERNANCE board. All STAKEHOLDERS are encouraged to provide '
        'feedback through the designated channels.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
