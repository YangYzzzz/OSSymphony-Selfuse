"""
Reward Script: Add 1.27 cm left indent to all body text paragraphs in research proposal
Task ID: writer_para_004
Domain: libreoffice_writer
Scoring:
  Component 1: Intro body paragraph (para index 2) has left_indent = 457200 EMU (~1.27 cm)  — 0.25 pts
  Component 2: Objectives body paragraph (para index 4) has left_indent = 457200 EMU         — 0.25 pts
  Component 3: Methodology body paragraph (para index 6) has left_indent = 457200 EMU        — 0.25 pts
  Component 4: All 3 body paragraphs correctly indented AND no heading para has indent > 0   — 0.25 pts
  Total: 1.0
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_para_004'

# Target indent value: 1.27 cm = 0.5 inches = 457200 EMU
TARGET_INDENT_EMU = 457200
# Tolerance: ±5000 EMU (~0.014 cm) to account for minor rounding
INDENT_TOLERANCE = 5000

# Expected paragraph indices and their roles (0-indexed)
# Para 0: title (heading)    — no indent expected
# Para 1: Introduction       — no indent expected (section heading)
# Para 2: intro body text    — MUST have 1.27 cm indent
# Para 3: Objectives         — no indent expected (section heading)
# Para 4: objectives body text — MUST have 1.27 cm indent
# Para 5: Methodology        — no indent expected (section heading)
# Para 6: methodology body text — MUST have 1.27 cm indent

BODY_PARA_INDICES = [2, 4, 6]
HEADING_PARA_INDICES = [0, 1, 3, 5]

# Partial text signatures to identify paragraphs (for robustness)
BODY_PARA_SIGNATURES = {
    2: 'This proposal outlines',
    4: 'The primary objective',
    6: 'We will employ a combination',
}
HEADING_PARA_SIGNATURES = {
    0: 'Research Proposal',
    1: 'Introduction',
    3: 'Objectives',
    5: 'Methodology',
}


def is_indent_correct(indent_emu):
    """Check if indent value is ~1.27 cm (457200 EMU) within tolerance."""
    if indent_emu is None:
        return False
    return abs(indent_emu - TARGET_INDENT_EMU) <= INDENT_TOLERANCE


def is_no_indent(indent_emu):
    """Check if indent is 0 or None (no indent set)."""
    if indent_emu is None:
        return True
    return indent_emu == 0 or abs(indent_emu) <= INDENT_TOLERANCE


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Task: Add 1.27 cm left indent to body text paragraphs 2, 4, 6.
    Heading paragraphs 0, 1, 3, 5 must remain without indentation.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 7 paragraphs
    if len(doc.paragraphs) < 7:
        print(f"CRITICAL: Expected >=7 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    paragraphs = doc.paragraphs

    # Verify paragraph identity (text content should match expected)
    for idx, sig in BODY_PARA_SIGNATURES.items():
        if idx < len(paragraphs) and not paragraphs[idx].text.startswith(sig):
            print(f"WARN: Para {idx} text '{paragraphs[idx].text[:40]}' doesn't start with '{sig}'")

    # Component 1: Introduction body paragraph (para 2) has left_indent ~1.27 cm (0.25 points)
    try:
        para2 = paragraphs[2]
        indent2 = para2.paragraph_format.left_indent
        if is_indent_correct(indent2):
            print(f"PASS: Component 1 — Para 2 (intro body) left_indent={indent2} EMU (~1.27 cm) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Para 2 (intro body) left_indent={indent2} EMU, expected ~{TARGET_INDENT_EMU} EMU (1.27 cm)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Objectives body paragraph (para 4) has left_indent ~1.27 cm (0.25 points)
    try:
        para4 = paragraphs[4]
        indent4 = para4.paragraph_format.left_indent
        if is_indent_correct(indent4):
            print(f"PASS: Component 2 — Para 4 (objectives body) left_indent={indent4} EMU (~1.27 cm) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — Para 4 (objectives body) left_indent={indent4} EMU, expected ~{TARGET_INDENT_EMU} EMU (1.27 cm)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Methodology body paragraph (para 6) has left_indent ~1.27 cm (0.25 points)
    try:
        para6 = paragraphs[6]
        indent6 = para6.paragraph_format.left_indent
        if is_indent_correct(indent6):
            print(f"PASS: Component 3 — Para 6 (methodology body) left_indent={indent6} EMU (~1.27 cm) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 3 — Para 6 (methodology body) left_indent={indent6} EMU, expected ~{TARGET_INDENT_EMU} EMU (1.27 cm)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: All 3 body paragraphs correctly indented AND no heading paragraph has indent > 0
    # This is a compound check: body indents must be set (ensures task was done) AND headings must not have been
    # accidentally indented. Fails on initial because body paragraphs are not indented (left_indent=0).
    try:
        body_all_correct = all(
            is_indent_correct(paragraphs[idx].paragraph_format.left_indent)
            for idx in BODY_PARA_INDICES
            if idx < len(paragraphs)
        )
        heading_no_indent = all(
            is_no_indent(paragraphs[idx].paragraph_format.left_indent)
            for idx in HEADING_PARA_INDICES
            if idx < len(paragraphs)
        )
        heading_indents = {
            idx: paragraphs[idx].paragraph_format.left_indent
            for idx in HEADING_PARA_INDICES
            if idx < len(paragraphs)
        }
        if body_all_correct and heading_no_indent:
            print(f"PASS: Component 4 — All body paragraphs correctly indented AND no heading has indent > 0 (0.25 pts)")
            print(f"  Heading indents: {heading_indents}")
            total_score += 0.25
        elif not body_all_correct:
            print(f"FAIL: Component 4 — Not all body paragraphs are correctly indented (prerequisite for this compound check)")
        else:
            print(f"FAIL: Component 4 — One or more heading paragraphs has unexpected indent > 0: {heading_indents}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
