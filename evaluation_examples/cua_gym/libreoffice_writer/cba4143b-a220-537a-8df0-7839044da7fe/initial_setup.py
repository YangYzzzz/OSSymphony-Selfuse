"""
Initial Setup: Research proposal document with no body text indentation
Task ID: writer_para_004
Domain: libreoffice_writer

Creates a research proposal document with 7 paragraphs:
- Para 1: Title (center-aligned, no indent)
- Para 2: 'Introduction' heading (no indent)
- Para 3: Body text (no indent — agent must add 1.27cm indent)
- Para 4: 'Objectives' heading (no indent)
- Para 5: Body text (no indent — agent must add 1.27cm indent)
- Para 6: 'Methodology' heading (no indent)
- Para 7: Body text (no indent — agent must add 1.27cm indent)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_para_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Title — center-aligned, no indent
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('Research Proposal: Machine Learning in Healthcare')
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Paragraph 2: 'Introduction' heading — no indent
    intro_heading = doc.add_paragraph()
    intro_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    intro_run = intro_heading.add_run('Introduction')
    intro_run.bold = True
    intro_run.font.size = Pt(14)

    # Paragraph 3: Body text — NO indent (agent task: add 1.27 cm indent)
    body1 = doc.add_paragraph()
    body1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # Explicitly ensure no left indent
    body1.paragraph_format.left_indent = Inches(0)
    body1.add_run(
        'This proposal outlines a three-year research program aimed at developing '
        'machine learning models for early disease detection using electronic health records.'
    )

    # Paragraph 4: 'Objectives' heading — no indent
    obj_heading = doc.add_paragraph()
    obj_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    obj_run = obj_heading.add_run('Objectives')
    obj_run.bold = True
    obj_run.font.size = Pt(14)

    # Paragraph 5: Body text — NO indent (agent task: add 1.27 cm indent)
    body2 = doc.add_paragraph()
    body2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    body2.paragraph_format.left_indent = Inches(0)
    body2.add_run(
        'The primary objective is to create robust predictive models that can identify '
        'patients at risk of chronic diseases at least six months before clinical diagnosis.'
    )

    # Paragraph 6: 'Methodology' heading — no indent
    meth_heading = doc.add_paragraph()
    meth_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    meth_run = meth_heading.add_run('Methodology')
    meth_run.bold = True
    meth_run.font.size = Pt(14)

    # Paragraph 7: Body text — NO indent (agent task: add 1.27 cm indent)
    body3 = doc.add_paragraph()
    body3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    body3.paragraph_format.left_indent = Inches(0)
    body3.add_run(
        'We will employ a combination of supervised and unsupervised learning techniques on '
        'anonymized datasets from three partner hospitals comprising over 2 million patient records.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
