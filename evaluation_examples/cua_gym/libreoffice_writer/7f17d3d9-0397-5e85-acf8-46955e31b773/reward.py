"""
Reward Script: Set up document structure with bookmarks, TOC, heading numbering, and document properties.
Task ID: writer_struct_080
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Three bookmarks 'intro', 'analysis', 'summary' at the correct headings
  Component 2 (0.25): TOC field covering levels 1-2 present at the beginning of document
  Component 3 (0.25): Chapter numbering applied to all Heading 1 and Heading 2 paragraphs
  Component 4 (0.25): Document properties — title, author, keywords set correctly
"""

import os
import re

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_080'
FILE_PATH = f'{WORKDIR}/performance_review.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------------------------------------------------------------
    # Component 1: Three bookmarks at correct headings (0.25 points)
    # Task requires: 'intro' at 'Introduction', 'analysis' at 'Analysis',
    #               'summary' at 'Summary'
    # ---------------------------------------------------------------
    try:
        # Collect (bookmark_name -> paragraph_text) mapping
        bookmark_locations = {}
        for para in doc.paragraphs:
            xml = para._element.xml
            bm_starts = re.findall(r'<w:bookmarkStart[^>]*w:name="([^"]+)"', xml)
            for bm_name in bm_starts:
                bookmark_locations[bm_name] = (para.text, getattr(para.style, 'name', ''))

        # Check required bookmarks exist and are placed at the correct headings
        required = {
            'intro': 'Introduction',
            'analysis': 'Analysis',
            'summary': 'Summary',
        }
        bm_pass_count = 0
        for bm_name, expected_heading in required.items():
            if bm_name not in bookmark_locations:
                print(f"FAIL: Component 1 — bookmark '{bm_name}' not found "
                      f"(expected at heading '{expected_heading}')")
            else:
                actual_text, actual_style = bookmark_locations[bm_name]
                if expected_heading.lower() not in actual_text.lower():
                    print(f"FAIL: Component 1 — bookmark '{bm_name}' found at paragraph '{actual_text}' "
                          f"(expected at '{expected_heading}')")
                else:
                    print(f"PASS: Component 1 — bookmark '{bm_name}' correctly placed at '{actual_text}' "
                          f"(style: {actual_style})")
                    bm_pass_count += 1

        if bm_pass_count == len(required):
            print("PASS: Component 1 — all 3 bookmarks (intro, analysis, summary) correctly placed (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — only {bm_pass_count}/{len(required)} bookmarks correctly placed (0.0 pts)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------------------------------------------------------------
    # Component 2: TOC field covering levels 1-2 at beginning (0.25 points)
    # Task requires: TOC at the beginning including levels 1-2
    # ---------------------------------------------------------------
    try:
        toc_para_index = None
        toc_field_str = None

        for i, para in enumerate(doc.paragraphs):
            xml = para._element.xml
            if 'instrText' in xml:
                matches = re.findall(r'<w:instrText[^>]*>([^<]+)</w:instrText>', xml)
                for m in matches:
                    if 'TOC' in m.upper():
                        toc_para_index = i
                        toc_field_str = m.strip()
                        print(f"INFO: TOC field found at paragraph [{i}]: '{toc_field_str}'")
                        break
            if toc_para_index is not None:
                break

        if toc_para_index is None:
            print("FAIL: Component 2 — No TOC field found in document")
        elif toc_para_index > 10:
            print(f"FAIL: Component 2 — TOC field found at paragraph [{toc_para_index}] but not at beginning (>10)")
        else:
            # Check level range: \o "1-2" or broader (1-3 also covers 1-2), or no \o (all levels)
            level_check = (
                '\\o "1-2"' in toc_field_str or
                "\\o '1-2'" in toc_field_str or
                re.search(r'\\o\s+["\']1-[2-9]["\']', toc_field_str) is not None or
                '\\o' not in toc_field_str
            )
            if level_check:
                print(f"PASS: Component 2 — TOC field covers levels 1-2 (field: {toc_field_str!r})")
                print("PASS: Component 2 — TOC field present at beginning of document (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — TOC field found but level range may not include 1-2: "
                      f"{toc_field_str!r}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---------------------------------------------------------------
    # Component 3: Chapter numbering applied to Heading 1 and Heading 2 (0.25 points)
    # Task requires: 1., 1.1. style numbering on headings
    # ---------------------------------------------------------------
    try:
        h1_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Heading 1']
        h2_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Heading 2']

        h1_with_num = 0
        h2_with_num = 0
        num_id_used = None

        for para in h1_paras:
            xml = para._element.xml
            numid_match = re.search(r'<w:numId w:val="(\d+)"', xml)
            ilvl_match = re.search(r'<w:ilvl w:val="(\d+)"', xml)
            if numid_match and ilvl_match:
                h1_with_num += 1
                num_id_used = numid_match.group(1)

        for para in h2_paras:
            xml = para._element.xml
            numid_match = re.search(r'<w:numId w:val="(\d+)"', xml)
            ilvl_match = re.search(r'<w:ilvl w:val="(\d+)"', xml)
            if numid_match and ilvl_match:
                h2_with_num += 1

        print(f"INFO: H1 paragraphs total={len(h1_paras)}, with numbering={h1_with_num}")
        print(f"INFO: H2 paragraphs total={len(h2_paras)}, with numbering={h2_with_num}")

        # Verify numbering definition uses decimal outline (1., 1.1., etc.)
        # Returns True if the numbering format matches 1./1.1. pattern, False otherwise
        def _check_numbering_format(num_id, numbering_part):
            if not num_id or not numbering_part:
                return False
            num_xml = numbering_part.element.xml
            abstract_id_match = re.search(
                rf'<w:num w:numId="{num_id}">.*?<w:abstractNumId w:val="(\d+)"',
                num_xml, re.DOTALL
            )
            if not abstract_id_match:
                return False
            abstract_id = abstract_id_match.group(1)
            abstract_match = re.search(
                rf'<w:abstractNum w:abstractNumId="{abstract_id}">(.*?)</w:abstractNum>',
                num_xml, re.DOTALL
            )
            if not abstract_match:
                return False
            content = abstract_match.group(0)
            lvl_texts = re.findall(r'<w:lvlText w:val="([^"]+)"', content)
            print(f"INFO: Numbering level texts: {lvl_texts}")
            has_l1_fmt = any(re.match(r'%1\.?$', t) for t in lvl_texts)
            has_l2_fmt = any(re.match(r'%1\.%2\.?', t) for t in lvl_texts)
            return has_l1_fmt and has_l2_fmt

        numbering_format_ok = _check_numbering_format(num_id_used, doc.part.numbering_part)
        if numbering_format_ok:
            print("PASS: Component 3 — chapter numbering format verified as 1./1.1. pattern")
        else:
            print("FAIL: Component 3 — numbering format not verified as 1./1.1. pattern")

        all_h1_numbered = (h1_with_num == len(h1_paras) and len(h1_paras) > 0)
        all_h2_numbered = (h2_with_num == len(h2_paras) and len(h2_paras) > 0)

        if all_h1_numbered and all_h2_numbered and numbering_format_ok:
            print("PASS: Component 3 — chapter numbering (1., 1.1.) applied to all headings (0.25 pts)")
            total_score += 0.25
        else:
            if not all_h1_numbered:
                print(f"FAIL: Component 3 — {h1_with_num}/{len(h1_paras)} Heading 1 paragraphs have numbering")
            if not all_h2_numbered:
                print(f"FAIL: Component 3 — {h2_with_num}/{len(h2_paras)} Heading 2 paragraphs have numbering")
            if not numbering_format_ok:
                print("FAIL: Component 3 — numbering format not verified as 1./1.1. pattern")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ---------------------------------------------------------------
    # Component 4: Document properties set correctly (0.25 points)
    # Task requires: title='Annual Performance Review', author='HR Department',
    #               keywords='performance, annual, review'
    # ---------------------------------------------------------------
    try:
        cp = doc.core_properties
        actual_title = cp.title or ''
        actual_author = cp.author or ''
        actual_keywords = cp.keywords or ''

        print(f"INFO: Title='{actual_title}', Author='{actual_author}', Keywords='{actual_keywords}'")

        title_ok = (actual_title.strip() == 'Annual Performance Review')
        author_ok = (actual_author.strip() == 'HR Department')

        # Keywords: normalize — split on comma/semicolon, compare required set
        expected_kw = {'performance', 'annual', 'review'}
        actual_kw = {k.strip().lower() for k in re.split(r'[,;]', actual_keywords) if k.strip()}
        keywords_ok = (expected_kw <= actual_kw)

        props_fail_count = 0
        if not title_ok:
            print(f"FAIL: Component 4 — title mismatch: got '{actual_title}', expected 'Annual Performance Review'")
            props_fail_count += 1
        else:
            print(f"PASS: Component 4 — title: '{actual_title}'")

        if not author_ok:
            print(f"FAIL: Component 4 — author mismatch: got '{actual_author}', expected 'HR Department'")
            props_fail_count += 1
        else:
            print(f"PASS: Component 4 — author: '{actual_author}'")

        if not keywords_ok:
            print(f"FAIL: Component 4 — keywords mismatch: got '{actual_keywords}', "
                  f"expected to include 'performance, annual, review'")
            props_fail_count += 1
        else:
            print(f"PASS: Component 4 — keywords: '{actual_keywords}'")

        if props_fail_count == 0:
            print("PASS: Component 4 — document properties correctly set (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 4 — {props_fail_count} property check(s) failed")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
