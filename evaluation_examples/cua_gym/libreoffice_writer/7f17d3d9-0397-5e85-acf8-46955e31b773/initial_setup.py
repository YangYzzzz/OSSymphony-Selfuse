"""
Initial Setup: Annual Performance Review document without bookmarks, TOC, heading numbering, or document properties.
Task ID: writer_struct_080
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'performance_review'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # ----- Document Title -----
    title_para = doc.add_paragraph()
    title_para.style = doc.styles['Title']
    title_run = title_para.add_run('Annual Performance Review 2025')
    title_run.font.size = Pt(24)

    subtitle_para = doc.add_paragraph()
    subtitle_para.style = doc.styles['Subtitle']
    subtitle_para.add_run('Fiscal Year January – December 2025')

    doc.add_paragraph()

    # ==========================================
    # Heading 1: Introduction
    # ==========================================
    doc.add_heading('Introduction', level=1)

    doc.add_paragraph(
        'This document presents the comprehensive annual performance review for fiscal year 2025. '
        'The review encompasses all business units, departments, and key operational metrics that '
        'define organizational health and strategic progress. Leadership teams, department heads, '
        'and HR professionals use this review as a foundation for strategic planning and talent development.'
    )
    doc.add_paragraph(
        'The evaluation framework follows the established KPI methodology adopted in 2023, '
        'incorporating both quantitative performance indicators and qualitative assessments from '
        'peer reviews, manager evaluations, and self-assessments.'
    )

    # H2 under Introduction
    doc.add_heading('Background and Methodology', level=2)
    doc.add_paragraph(
        'Our performance evaluation methodology draws on industry best practices from leading '
        'HR research institutions. Data was collected across all 14 regional offices between '
        'October and November 2025, with additional input from the global workforce survey '
        'conducted in September. A total of 2,847 employees participated across all levels.'
    )
    doc.add_paragraph(
        'Qualitative data was gathered through structured interviews, 360-degree feedback '
        'sessions, and manager assessments. Quantitative metrics were pulled directly from '
        'the HRIS system and verified by the Finance team.'
    )

    doc.add_heading('Scope and Coverage', level=2)
    doc.add_paragraph(
        'This review covers all permanent employees across the following divisions: Engineering, '
        'Marketing, Sales, Operations, Finance, Human Resources, Legal, and Customer Success. '
        'Contract workers and part-time employees are covered in a separate supplementary report.'
    )

    doc.add_paragraph(
        'Geographic coverage includes North America (48%), Europe (29%), Asia-Pacific (18%), '
        'and Rest of World (5%). All figures are presented in USD unless otherwise stated.'
    )

    doc.add_page_break()

    # ==========================================
    # Heading 1: Performance Metrics
    # ==========================================
    doc.add_heading('Performance Metrics', level=1)

    doc.add_paragraph(
        'The core performance metrics for 2025 reflect strong organizational momentum across '
        'most business units. Revenue per employee increased by 12.4% year-over-year, reaching '
        '$187,400 compared to $166,700 in 2024. Customer satisfaction scores averaged 4.6/5.0 '
        'across all product lines, representing a 0.3-point improvement.'
    )

    doc.add_heading('Financial Performance Indicators', level=2)
    doc.add_paragraph(
        'Total revenue for FY2025 reached $532.4 million, representing an 18.7% increase over '
        'FY2024. Gross margin improved to 67.3% from 64.1%, driven by operational efficiencies '
        'in the Engineering and Operations divisions. EBITDA margin expanded to 23.1%, up from '
        '19.8% in the prior year.'
    )

    # Table: Key Financial Metrics
    table1 = doc.add_table(rows=6, cols=3)
    table1.style = 'Table Grid'
    headers = ['Metric', 'FY2024', 'FY2025']
    for col_idx, header in enumerate(headers):
        cell = table1.cell(0, col_idx)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
    data_rows = [
        ('Total Revenue', '$448.3M', '$532.4M'),
        ('Gross Margin', '64.1%', '67.3%'),
        ('EBITDA Margin', '19.8%', '23.1%'),
        ('Net Income', '$68.2M', '$89.7M'),
        ('Revenue per Employee', '$166,700', '$187,400'),
    ]
    for r_idx, (m, v24, v25) in enumerate(data_rows, 1):
        table1.cell(r_idx, 0).text = m
        table1.cell(r_idx, 1).text = v24
        table1.cell(r_idx, 2).text = v25

    doc.add_paragraph()

    doc.add_heading('Operational Performance Indicators', level=2)
    doc.add_paragraph(
        'Operational efficiency metrics showed consistent improvement throughout the year. '
        'Average project delivery time decreased by 8.2 days compared to 2024, while defect '
        'rates in the Engineering division dropped to 1.2% from 2.7%. Customer support ticket '
        'resolution time improved by 31%, averaging 4.2 hours per ticket.'
    )
    doc.add_paragraph(
        'Supply chain disruptions were minimized through the new vendor diversification '
        'program launched in Q1, reducing single-source dependency from 34% to 11%. '
        'Inventory turnover improved to 8.4x annually.'
    )

    doc.add_heading('Employee Performance Distribution', level=2)
    doc.add_paragraph(
        'Performance ratings followed a near-normal distribution with slight right-skew. '
        'Exceptional performers (rating 5/5) comprised 11.2% of the workforce, up from 8.7% '
        'in 2024. High performers (rating 4/5) represented 37.4%, while meeting-expectations '
        '(rating 3/5) was the modal category at 41.8%.'
    )

    doc.add_page_break()

    # ==========================================
    # Heading 1: Analysis
    # ==========================================
    doc.add_heading('Analysis', level=1)

    doc.add_paragraph(
        'The performance data reveals several meaningful trends that warrant leadership attention '
        'in strategic planning for 2026. Cross-functional collaboration has emerged as a key '
        'differentiator for high-performing teams, while technical skill gaps remain the primary '
        'constraint on growth in the Engineering division.'
    )
    doc.add_paragraph(
        'Regression analysis of the performance data indicates a strong correlation (r=0.71) '
        'between manager effectiveness scores and team-level performance outcomes, reinforcing '
        'the case for continued investment in management development programs.'
    )

    doc.add_heading('Departmental Analysis', level=2)
    doc.add_paragraph(
        'The Sales department achieved 108% of its annual quota, with the Enterprise segment '
        'outperforming at 121%. The mid-market segment fell short at 94%, primarily due to '
        'two high-performing reps departing in Q2. Compensation adjustments and targeted '
        'recruitment should address this gap in H1 2026.'
    )
    doc.add_paragraph(
        'Marketing demonstrated strong ROI on campaign spend, achieving a 4.8x return on '
        'digital advertising investment. Content marketing drove 34% of new pipeline, up from '
        '22% in 2024. Event marketing ROI declined as virtual-to-in-person transition costs '
        'were higher than budgeted.'
    )

    doc.add_heading('Talent Retention Analysis', level=2)
    doc.add_paragraph(
        'Overall voluntary attrition rate was 11.3%, slightly below the industry average of '
        '12.8% for technology companies. Engineering experienced the highest attrition at 14.2%, '
        'while Customer Success had the lowest at 7.6%. Exit interviews cited compensation '
        'competitiveness and career advancement as the top two factors for departures.'
    )
    doc.add_paragraph(
        'The new Career Development Framework launched in Q2 showed early positive results, '
        'with promoted employees reporting 23% higher engagement scores in the Q4 pulse survey. '
        'Internal mobility rate increased to 15.4%, the highest in company history.'
    )

    doc.add_heading('Diversity and Inclusion Metrics', level=2)
    doc.add_paragraph(
        'The company made measurable progress on D&I commitments in 2025. Women in senior '
        'leadership roles increased from 31% to 37%, on track toward the 40% target by 2027. '
        'Underrepresented minorities in technical roles increased from 18% to 22%. Pay equity '
        'analysis confirmed <2% unexplained variance in compensation across demographic groups.'
    )

    doc.add_page_break()

    # ==========================================
    # Heading 1: Summary
    # ==========================================
    doc.add_heading('Summary', level=1)

    doc.add_paragraph(
        'Fiscal year 2025 represents a strong performance year for the organization across '
        'nearly all measured dimensions. Financial results exceeded targets, operational '
        'efficiency improved materially, and employee engagement reached an all-time high '
        'of 78% favorable. The investments made in talent development, technology infrastructure, '
        'and market expansion are clearly yielding returns.'
    )
    doc.add_paragraph(
        'Three areas require prioritized attention in 2026: Engineering talent retention and '
        'compensation competitiveness, mid-market sales capacity restoration, and the continued '
        'scaling of the Career Development Framework across all divisions.'
    )

    doc.add_heading('Key Achievements', level=2)
    achievements = [
        'Revenue growth of 18.7% YoY, exceeding the 15% target',
        'EBITDA margin expansion of 3.3 percentage points',
        'Employee engagement score of 78%, highest in company history',
        'Voluntary attrition below industry average for third consecutive year',
        'D&I progress across all tracked demographic dimensions',
        'Successful launch of Career Development Framework with measurable impact',
        'Supply chain risk reduction from 34% to 11% single-source dependency',
    ]
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')

    doc.add_heading('Areas for Improvement', level=2)
    doc.add_paragraph(
        'While the overall performance trajectory is positive, three strategic gaps require '
        'leadership focus. The Engineering division talent pipeline needs strengthening through '
        'partnerships with universities and a revised total compensation strategy. Mid-market '
        'sales recovery depends on successful Q1 2026 hiring execution. The event marketing '
        'budget model needs recalibration for the hybrid event environment.'
    )

    doc.add_page_break()

    # ==========================================
    # Heading 1: Recommendations
    # ==========================================
    doc.add_heading('Recommendations', level=1)

    doc.add_paragraph(
        'Based on the performance data and analysis presented in this review, the HR Leadership '
        'Team makes the following strategic recommendations for fiscal year 2026 planning and '
        'execution. These recommendations are prioritized by impact potential and implementation '
        'feasibility.'
    )

    doc.add_heading('Talent Strategy Recommendations', level=2)
    doc.add_paragraph(
        'Immediately benchmark Engineering compensation against the 75th percentile of market '
        'data and adjust base salaries accordingly. Establish university partnerships with at '
        'least 5 target institutions for 2026 graduate hiring. Expand the internal mobility '
        'program to facilitate cross-functional transfers and reduce external hiring costs.'
    )
    recommendations_list = [
        'Conduct compensation benchmarking review by end of Q1 2026',
        'Launch university partnership program with target institutions',
        'Expand Career Development Framework to all departments by Q2 2026',
        'Implement structured mentoring program for high-potential employees',
        'Introduce retention bonuses for critical technical roles',
    ]
    for rec in recommendations_list:
        doc.add_paragraph(rec, style='List Number')

    doc.add_heading('Operational Recommendations', level=2)
    doc.add_paragraph(
        'Continue investment in operational automation to sustain the efficiency gains realized '
        'in 2025. The project management transformation initiative should be extended to the '
        'Marketing and Sales divisions. Customer support operations should scale headcount '
        'proportionally with anticipated revenue growth of 20-25% in 2026.'
    )
    doc.add_paragraph(
        'The vendor diversification program should be maintained and extended to software '
        'vendors where over-dependence on single providers creates operational risk. '
        'Quarterly business reviews with strategic vendors should be institutionalized.'
    )

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
