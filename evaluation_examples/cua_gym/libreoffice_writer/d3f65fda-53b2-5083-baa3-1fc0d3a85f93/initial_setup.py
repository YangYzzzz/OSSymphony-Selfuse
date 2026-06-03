"""
Initial Setup: Legal analysis document with 4 sections and one existing endnote.
Task ID: osworld_writer_bibliography_crossref_004
Domain: libreoffice_writer

Creates a .docx with 4 sections, one pre-existing endnote, and the target sentence
'The regulation was first introduced in 1998.' in Section 2 WITHOUT the new endnote.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree
import copy
import zipfile
import shutil

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_bibliography_crossref_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_endnote_support(doc):
    """
    Set up the endnotes part in the document if not already present,
    and return the endnotes element.
    The endnotes.xml must have the standard separator and continuationSeparator notes (id=0, id=1).
    """
    # Access or create the endnotes part
    part = doc.part

    # Check if endnotes part already exists
    endnotes_part = None
    for rel in part.rels.values():
        if 'endnotes' in rel.reltype:
            endnotes_part = rel.target_part
            break

    if endnotes_part is None:
        # Need to create endnotes.xml from scratch
        # We'll do this by manipulating the docx file directly after saving
        pass

    return endnotes_part


def create_endnote_xml_with_notes(notes_data):
    """
    Create the endnotes.xml content with:
      - id=0: separator
      - id=1: continuationSeparator
      - id=2+: actual endnotes
    notes_data: list of (id, text) for actual endnotes
    """
    nsmap = {
        'wpc': 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas',
        'cx': 'http://schemas.microsoft.com/office/drawing/2014/chartex',
        'cx1': 'http://schemas.microsoft.com/office/drawing/2015/9/8/chartex',
        'cx2': 'http://schemas.microsoft.com/office/drawing/2015/10/21/chartex',
        'cx3': 'http://schemas.microsoft.com/office/drawing/2016/5/9/chartex',
        'cx4': 'http://schemas.microsoft.com/office/drawing/2016/5/10/chartex',
        'cx5': 'http://schemas.microsoft.com/office/drawing/2016/5/11/chartex',
        'cx6': 'http://schemas.microsoft.com/office/drawing/2016/5/12/chartex',
        'cx7': 'http://schemas.microsoft.com/office/drawing/2016/5/13/chartex',
        'cx8': 'http://schemas.microsoft.com/office/drawing/2016/5/14/chartex',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'aink': 'http://schemas.microsoft.com/office/drawing/2016/ink',
        'am3d': 'http://schemas.microsoft.com/office/drawing/2017/model3d',
        'o': 'urn:schemas-microsoft-com:office:office',
        'oel': 'http://schemas.microsoft.com/office/2019/extlst',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
        'v': 'urn:schemas-microsoft-com:vml',
        'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'w10': 'urn:schemas-microsoft-com:office:word',
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
        'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
        'w16cex': 'http://schemas.microsoft.com/office/word/2018/wordml/cex',
        'w16cid': 'http://schemas.microsoft.com/office/word/2016/wordml/cid',
        'w16': 'http://schemas.microsoft.com/office/word/2018/wordml',
        'w16sdtdh': 'http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash',
        'w16se': 'http://schemas.microsoft.com/office/word/2015/wordml/symex',
        'wpg': 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup',
        'wpi': 'http://schemas.microsoft.com/office/word/2010/wordprocessingInk',
        'wne': 'http://schemas.microsoft.com/office/word/2006/wordml',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    }

    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    xml_lines = []
    xml_lines.append('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>')
    xml_lines.append('<w:endnotes xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex" xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid" xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml" xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" mc:Ignorable="w14 w15 w16se w16cid w16 w16cex wpg wpi wne wps wp14">')

    # Separator (id=0)
    xml_lines.append('  <w:endnote w:type="separator" w:id="0"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p></w:endnote>')
    # Continuation separator (id=1)
    xml_lines.append('  <w:endnote w:type="continuationSeparator" w:id="1"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p></w:endnote>')

    # Actual endnotes
    for note_id, note_text in notes_data:
        xml_lines.append(f'  <w:endnote w:id="{note_id}"><w:p><w:pPr><w:pStyle w:val="EndnoteText"/></w:pPr><w:r><w:rPr><w:rStyle w:val="EndnoteReference"/></w:rPr><w:endnoteRef/></w:r><w:r><w:t xml:space="preserve"> {note_text}</w:t></w:r></w:p></w:endnote>')

    xml_lines.append('</w:endnotes>')
    return '\n'.join(xml_lines)


def build_initial_docx():
    """Build the initial .docx with 4 sections and one pre-existing endnote."""
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title
    title = doc.add_heading('Legal Analysis: Regulatory Compliance in the Digital Age', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Author/Date line
    meta = doc.add_paragraph('Prepared by: Legal Research Division | Date: March 2025')
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta.runs[0].font.italic = True
    meta.runs[0].font.size = Pt(11)

    doc.add_paragraph('')  # spacer

    # ------ SECTION 1 ------
    doc.add_heading('Section 1: Introduction to Regulatory Compliance', level=1)

    p1 = doc.add_paragraph(
        'Regulatory compliance has become a cornerstone of modern business operations. '
        'Organizations across industries must navigate a complex web of laws and standards '
        'that govern their activities. Failure to comply can result in significant financial '
        'penalties, reputational damage, and in severe cases, criminal liability.'
    )

    p2 = doc.add_paragraph(
        'This analysis examines the evolution of regulatory frameworks in the context of '
        'digital transformation. The intersection of technology and law has created new '
        'challenges for compliance officers and legal practitioners alike.'
    )

    # Paragraph with existing endnote reference
    # We'll add endnote reference inline using XML
    p3 = doc.add_paragraph()
    run_text = p3.add_run(
        'The foundational principles of regulatory compliance were established through '
        'a series of landmark legislative acts in the late twentieth century.'
    )
    # Add endnote reference mark (id=2) after this sentence
    endnote_ref_run = p3.add_run()
    endnote_ref_elem = OxmlElement('w:endnoteReference')
    endnote_ref_elem.set(qn('w:id'), '2')
    endnote_ref_run._element.append(endnote_ref_elem)
    run_after = p3.add_run(
        ' These principles have since been codified and expanded upon by successive '
        'governments and international bodies.'
    )

    p4 = doc.add_paragraph(
        'Understanding the historical context of these regulations is essential for any '
        'practitioner operating in today\'s legal environment. The following sections '
        'provide a comprehensive analysis of key regulatory domains.'
    )

    # ------ SECTION 2 ------
    doc.add_heading('Section 2: Historical Development of Regulatory Standards', level=1)

    p5 = doc.add_paragraph(
        'The development of modern regulatory standards spans several decades and multiple '
        'jurisdictions. Early regulatory efforts were fragmented and often reactive in '
        'nature, responding to specific incidents rather than proactively addressing '
        'systemic risks.'
    )

    # TARGET sentence — NO endnote here in initial state
    p6 = doc.add_paragraph(
        'The regulation was first introduced in 1998. '
        'This landmark legislation marked a significant shift in how governments approached '
        'the oversight of commercial activities, particularly in sectors deemed critical '
        'to national interest.'
    )

    p7 = doc.add_paragraph(
        'Subsequent amendments in 2003 and 2007 expanded the scope of the original '
        'regulation, incorporating provisions for digital records and electronic commerce. '
        'By 2010, the framework had been adopted by over forty countries, establishing '
        'it as a de facto international standard.'
    )

    p8 = doc.add_paragraph(
        'The enforcement mechanisms introduced during this period proved instrumental '
        'in shaping corporate behavior. Companies that previously operated in regulatory '
        'gray areas were compelled to establish formal compliance programs, leading to '
        'the professionalization of the compliance function.'
    )

    # ------ SECTION 3 ------
    doc.add_heading('Section 3: Current Regulatory Landscape', level=1)

    p9 = doc.add_paragraph(
        'Today\'s regulatory environment is characterized by increasing complexity and '
        'cross-jurisdictional reach. The globalization of business operations has made '
        'it necessary for organizations to maintain compliance with multiple overlapping '
        'regulatory regimes simultaneously.'
    )

    p10 = doc.add_paragraph(
        'Digital technologies have both complicated and facilitated compliance efforts. '
        'On one hand, the volume and velocity of regulated activities have increased '
        'exponentially. On the other hand, RegTech solutions now enable automated '
        'monitoring and reporting that was previously impossible.'
    )

    p11 = doc.add_paragraph(
        'Key sectors subject to intensive regulatory oversight include financial services, '
        'healthcare, environmental management, and data protection. Each domain has '
        'developed its own specialized compliance infrastructure, including dedicated '
        'regulatory bodies, professional associations, and certification programs.'
    )

    p12 = doc.add_paragraph(
        'The trend toward principles-based regulation, as opposed to prescriptive '
        'rule-based approaches, has shifted the burden of interpretation to regulated '
        'entities. This requires a sophisticated understanding of regulatory intent '
        'and a robust internal governance framework.'
    )

    # ------ SECTION 4 ------
    doc.add_heading('Section 4: Future Trends and Recommendations', level=1)

    p13 = doc.add_paragraph(
        'Looking ahead, several trends are likely to shape the regulatory compliance '
        'landscape in the coming decade. Artificial intelligence and machine learning '
        'are expected to play an increasingly important role in both regulation and '
        'compliance management.'
    )

    p14 = doc.add_paragraph(
        'Regulators are increasingly using data analytics to identify patterns of '
        'non-compliance and target enforcement resources more effectively. Organizations '
        'that invest in compliance technology now will be better positioned to meet '
        'future regulatory requirements.'
    )

    p15 = doc.add_paragraph(
        'International harmonization of regulatory standards remains a priority for '
        'multilateral bodies. The convergence of data protection regimes following '
        'the GDPR has demonstrated that meaningful international cooperation is possible, '
        'though significant challenges remain.'
    )

    p16 = doc.add_paragraph(
        'In conclusion, organizations must adopt a proactive and adaptive approach to '
        'regulatory compliance. This requires ongoing investment in people, processes, '
        'and technology, as well as active engagement with regulators and industry peers.'
    )

    # Save the document first
    doc.save(OUTPUT)
    print(f'Initial docx saved (without endnotes XML): {OUTPUT}')

    # Now inject endnotes.xml into the .docx zip
    inject_endnotes_into_docx(OUTPUT, notes_data=[(2, 'Corporate Governance Act, Consolidated Edition, Parliament of the United Kingdom, 1998.')])
    print(f'Endnotes injected into docx: {OUTPUT}')


def inject_endnotes_into_docx(docx_path, notes_data):
    """
    Inject endnotes.xml into the .docx file by manipulating the zip archive directly.
    notes_data: list of (id, text) tuples for actual endnotes
    """
    import tempfile

    endnotes_xml = create_endnote_xml_with_notes(notes_data)

    tmp_path = docx_path + '.tmp'

    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(tmp_path, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
            # Check if endnotes.xml already exists
            existing_files = zin.namelist()
            has_endnotes = 'word/endnotes.xml' in existing_files

            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename == 'word/_rels/document.xml.rels' and not has_endnotes:
                    # Add endnotes relationship
                    rels_xml = data.decode('utf-8')
                    # Insert before </Relationships>
                    endnote_rel = '<Relationship Id="rId99" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes" Target="endnotes.xml"/>'
                    rels_xml = rels_xml.replace('</Relationships>', endnote_rel + '</Relationships>')
                    data = rels_xml.encode('utf-8')

                elif item.filename == 'word/document.xml':
                    # Ensure endnoteReference elements are present in the body
                    # (already added via OxmlElement above, but verify no duplicates)
                    pass

                elif item.filename == 'word/endnotes.xml':
                    # Replace with our new version
                    data = endnotes_xml.encode('utf-8')

                zout.writestr(item, data)

            # Add endnotes.xml if it didn't exist
            if not has_endnotes:
                zout.writestr('word/endnotes.xml', endnotes_xml.encode('utf-8'))

    shutil.move(tmp_path, docx_path)
    print('Endnotes XML injected successfully.')


def create_initial():
    build_initial_docx()
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
