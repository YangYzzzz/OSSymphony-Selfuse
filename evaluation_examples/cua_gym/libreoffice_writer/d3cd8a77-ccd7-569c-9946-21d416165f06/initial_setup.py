"""
Initial Setup: Convert plain text outline into multi-level list
Task ID: writer_lec_005
Domain: libreoffice_writer

Creates a Writer document with a course outline in plain text.
Hierarchy is indicated by tab characters (indentation).
No list numbering is applied - that is the task for the agent.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading("Advanced Data Science and Machine Learning", level=1)

    doc.add_paragraph("")  # blank line

    # The outline content: plain text with tabs indicating hierarchy
    # Level 0 (main topics): no tab
    # Level 1 (subtopics): one tab
    # Level 2 (sub-subtopics): two tabs

    outline_items = [
        # Topic 1
        (0, "Foundations of Statistical Learning"),
        (1, "Probability Theory and Distributions"),
        (2, "Bayesian vs Frequentist Approaches"),
        (2, "Common Distribution Families"),
        (1, "Linear Algebra for Machine Learning"),
        (1, "Optimization Methods and Gradient Descent"),
        (2, "Stochastic Gradient Descent Variants"),
        (2, "Learning Rate Scheduling Strategies"),
        # Topic 2
        (0, "Supervised Learning Algorithms"),
        (1, "Regression Techniques"),
        (2, "Ridge and Lasso Regularization"),
        (2, "Polynomial Feature Engineering"),
        (1, "Classification Methods"),
        (1, "Ensemble Learning and Boosting"),
        (2, "Random Forest Hyperparameter Tuning"),
        (2, "XGBoost and LightGBM Comparison"),
        # Topic 3
        (0, "Deep Learning and Neural Networks"),
        (1, "Feedforward Network Architecture"),
        (2, "Activation Functions and Their Properties"),
        (2, "Batch Normalization Techniques"),
        (1, "Convolutional Neural Networks"),
        (1, "Recurrent Networks and Transformers"),
        (2, "Attention Mechanism Fundamentals"),
        (2, "Pre-training and Fine-tuning Workflows"),
    ]

    for level, text in outline_items:
        # Use tab characters to indicate indentation level
        prefix = "\t" * level
        para = doc.add_paragraph(prefix + text)
        # Set a normal style - no list styling
        para.style = doc.styles['Normal']
        # Set consistent font
        for run in para.runs:
            run.font.name = "Liberation Sans"
            run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
