"""
Reward Script: Convert plain text outline to multi-level numbered list
Task ID: writer_lec_005
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): All outline paragraphs have numbering applied
  Component 2 (0.3): Correct level assignments (ilvl 0/1/2)
  Component 3 (0.2): Numbering format is decimal multi-level (not bullet)
  Component 4 (0.2): Text content preserved correctly
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_005'

# Expected hierarchy: paragraph index -> expected ilvl
# Paras 0-1 are heading and blank, not part of the outline list
# Paras 2-25 are the outline items
EXPECTED_LEVELS = {
    2: 0,   # Foundations of Statistical Learning
    3: 1,   # Probability Theory and Distributions
    4: 2,   # Bayesian vs Frequentist Approaches
    5: 2,   # Common Distribution Families
    6: 1,   # Linear Algebra for Machine Learning
    7: 1,   # Optimization Methods and Gradient Descent
    8: 2,   # Stochastic Gradient Descent Variants
    9: 2,   # Learning Rate Scheduling Strategies
    10: 0,  # Supervised Learning Algorithms
    11: 1,  # Regression Techniques
    12: 2,  # Ridge and Lasso Regularization
    13: 2,  # Polynomial Feature Engineering
    14: 1,  # Classification Methods
    15: 1,  # Ensemble Learning and Boosting
    16: 2,  # Random Forest Hyperparameter Tuning
    17: 2,  # XGBoost and LightGBM Comparison
    18: 0,  # Deep Learning and Neural Networks
    19: 1,  # Feedforward Network Architecture
    20: 2,  # Activation Functions and Their Properties
    21: 2,  # Batch Normalization Techniques
    22: 1,  # Convolutional Neural Networks
    23: 1,  # Recurrent Networks and Transformers
    24: 2,  # Attention Mechanism Fundamentals
    25: 2,  # Pre-training and Fine-tuning Workflows
}

# Expected text content (stripped of tabs) for each outline paragraph
EXPECTED_TEXTS = {
    2: "Foundations of Statistical Learning",
    3: "Probability Theory and Distributions",
    4: "Bayesian vs Frequentist Approaches",
    5: "Common Distribution Families",
    6: "Linear Algebra for Machine Learning",
    7: "Optimization Methods and Gradient Descent",
    8: "Stochastic Gradient Descent Variants",
    9: "Learning Rate Scheduling Strategies",
    10: "Supervised Learning Algorithms",
    11: "Regression Techniques",
    12: "Ridge and Lasso Regularization",
    13: "Polynomial Feature Engineering",
    14: "Classification Methods",
    15: "Ensemble Learning and Boosting",
    16: "Random Forest Hyperparameter Tuning",
    17: "XGBoost and LightGBM Comparison",
    18: "Deep Learning and Neural Networks",
    19: "Feedforward Network Architecture",
    20: "Activation Functions and Their Properties",
    21: "Batch Normalization Techniques",
    22: "Convolutional Neural Networks",
    23: "Recurrent Networks and Transformers",
    24: "Attention Mechanism Fundamentals",
    25: "Pre-training and Fine-tuning Workflows",
}

NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def get_num_info(para):
    """Extract numId and ilvl from a paragraph's numPr element."""
    numPr_els = para._element.findall('.//w:numPr', NS)
    if not numPr_els:
        return None, None
    numId_el = para._element.findall('.//w:numId', NS)
    ilvl_el = para._element.findall('.//w:ilvl', NS)
    numId = numId_el[0].get(f'{{{NS["w"]}}}val') if numId_el else None
    ilvl = ilvl_el[0].get(f'{{{NS["w"]}}}val') if ilvl_el else None
    return numId, int(ilvl) if ilvl is not None else None


def get_abstract_num_format(doc, numId_val):
    """Get the numbering format info for a given numId by resolving to abstractNum."""
    try:
        numbering_part = doc.part.numbering_part
        numbering_el = numbering_part._element

        # Find the num element with the matching numId (last one wins if duplicates)
        target_abstract_id = None
        for num_el in numbering_el.findall('w:num', NS):
            nid = num_el.get(f'{{{NS["w"]}}}numId')
            if nid == str(numId_val):
                abs_ref = num_el.find('w:abstractNumId', NS)
                if abs_ref is not None:
                    target_abstract_id = abs_ref.get(f'{{{NS["w"]}}}val')

        if target_abstract_id is None:
            return None

        # Find the abstract numbering definition
        for abs_num in numbering_el.findall('w:abstractNum', NS):
            abs_id = abs_num.get(f'{{{NS["w"]}}}abstractNumId')
            if abs_id == target_abstract_id:
                levels = {}
                for lvl in abs_num.findall('w:lvl', NS):
                    ilvl = lvl.get(f'{{{NS["w"]}}}ilvl')
                    num_fmt_el = lvl.find('w:numFmt', NS)
                    lvl_text_el = lvl.find('w:lvlText', NS)
                    fmt = num_fmt_el.get(f'{{{NS["w"]}}}val') if num_fmt_el is not None else None
                    txt = lvl_text_el.get(f'{{{NS["w"]}}}val') if lvl_text_el is not None else None
                    levels[int(ilvl)] = {'numFmt': fmt, 'lvlText': txt}
                return levels
    except Exception:
        return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has at least 26 paragraphs
    if len(doc.paragraphs) < 26:
        print(f"CRITICAL: Expected at least 26 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All outline paragraphs have numbering applied (0.3 points)
    # This checks that numPr is present on each of the 24 outline paragraphs
    try:
        numbered_count = 0
        total_outline = len(EXPECTED_LEVELS)  # 24
        for idx in EXPECTED_LEVELS:
            para = doc.paragraphs[idx]
            numId, ilvl = get_num_info(para)
            if numId is not None:
                numbered_count += 1
            else:
                print(f"  DETAIL: Para {idx} ({para.text[:40]!r}) has no numbering")

        ratio = numbered_count / total_outline
        if ratio == 1.0:
            print(f"PASS: Component 1 - All {total_outline} outline paragraphs have numbering (0.3 pts)")
            total_score += 0.3
        elif ratio > 0:
            partial = round(0.3 * ratio, 2)
            print(f"PARTIAL: Component 1 - {numbered_count}/{total_outline} paragraphs numbered ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 - No outline paragraphs have numbering applied")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Correct level assignments (ilvl 0/1/2) (0.3 points)
    # Each paragraph must have the correct ilvl matching the original indentation
    try:
        correct_levels = 0
        total_outline = len(EXPECTED_LEVELS)
        for idx, expected_ilvl in EXPECTED_LEVELS.items():
            para = doc.paragraphs[idx]
            numId, actual_ilvl = get_num_info(para)
            if actual_ilvl == expected_ilvl:
                correct_levels += 1
            else:
                print(f"  DETAIL: Para {idx} expected ilvl={expected_ilvl}, got ilvl={actual_ilvl}")

        ratio = correct_levels / total_outline
        if ratio == 1.0:
            print(f"PASS: Component 2 - All {total_outline} paragraphs at correct list level (0.3 pts)")
            total_score += 0.3
        elif ratio > 0:
            partial = round(0.3 * ratio, 2)
            print(f"PARTIAL: Component 2 - {correct_levels}/{total_outline} paragraphs at correct level ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 - No paragraphs at correct list level")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Numbering format is decimal multi-level (0.2 points)
    # The numbering must use decimal format with multi-level text patterns
    # (not bullets, not single-level decimal)
    try:
        # Find the numId used by outline paragraphs
        numIds_used = set()
        for idx in EXPECTED_LEVELS:
            para = doc.paragraphs[idx]
            numId, _ = get_num_info(para)
            if numId is not None:
                numIds_used.add(numId)

        if not numIds_used:
            print(f"FAIL: Component 3 - No numbering found on outline paragraphs")
        else:
            format_score = 0.0
            for nid in numIds_used:
                levels = get_abstract_num_format(doc, nid)
                if levels is None:
                    print(f"  DETAIL: Could not resolve numbering format for numId={nid}")
                    continue

                # Check level 0 is decimal
                has_decimal_l0 = levels.get(0, {}).get('numFmt') == 'decimal'
                # Check at least 3 levels defined with decimal format
                decimal_levels = sum(1 for l in levels.values() if l.get('numFmt') == 'decimal')
                has_multilevel = decimal_levels >= 3
                # Check lvlText patterns include parent level references
                has_compound_text = False
                if 1 in levels and levels[1].get('lvlText'):
                    # Level 1 text should reference level 0 (e.g., %1.%2.)
                    if '%1' in levels[1]['lvlText'] and '%2' in levels[1]['lvlText']:
                        has_compound_text = True
                if 2 in levels and levels[2].get('lvlText'):
                    if '%1' in levels[2]['lvlText'] and '%3' in levels[2]['lvlText']:
                        has_compound_text = True

                if has_decimal_l0:
                    format_score += 0.07
                if has_multilevel:
                    format_score += 0.07
                if has_compound_text:
                    format_score += 0.06

                print(f"  DETAIL: numId={nid}: decimal_l0={has_decimal_l0}, multilevel={has_multilevel}, compound_text={has_compound_text}")

            format_score = min(format_score, 0.2)
            if format_score >= 0.2:
                print(f"PASS: Component 3 - Decimal multi-level numbering format confirmed (0.2 pts)")
                total_score += 0.2
            elif format_score > 0:
                print(f"PARTIAL: Component 3 - Partial format match ({format_score} pts)")
                total_score += format_score
            else:
                print(f"FAIL: Component 3 - Numbering format is not decimal multi-level")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Text content preserved AND numbering applied (0.2 points)
    # This is a compound check: text must be correct AND the paragraph must have numbering.
    # This ensures it only passes on the golden (where numbering is applied), not the initial.
    try:
        matching_count = 0
        total_outline = len(EXPECTED_TEXTS)
        for idx, expected_text in EXPECTED_TEXTS.items():
            para = doc.paragraphs[idx]
            actual_text = para.text.strip().lstrip('\t')
            numId, _ = get_num_info(para)
            if actual_text == expected_text and numId is not None:
                matching_count += 1
            elif numId is None:
                pass  # No numbering — will fail as expected on initial
            else:
                print(f"  DETAIL: Para {idx} text mismatch: expected {expected_text!r}, got {actual_text[:60]!r}")

        ratio = matching_count / total_outline
        if ratio == 1.0:
            print(f"PASS: Component 4 - All {total_outline} paragraphs have numbering AND correct text (0.2 pts)")
            total_score += 0.2
        elif ratio > 0.5:
            partial = round(0.2 * ratio, 2)
            print(f"PARTIAL: Component 4 - {matching_count}/{total_outline} paragraphs with numbering+text ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 - {matching_count}/{total_outline} paragraphs have both numbering and correct text")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
