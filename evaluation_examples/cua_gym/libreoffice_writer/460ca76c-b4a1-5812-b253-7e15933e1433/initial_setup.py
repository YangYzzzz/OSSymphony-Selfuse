"""
Initial Setup: Create a business letter with delivery address in default position
Task ID: writer_lec_062
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_062'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup: A4 ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # --- Sender Address (top right, small font) ---
    sender = doc.add_paragraph()
    sender.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    sender.paragraph_format.space_after = Pt(0)
    run = sender.add_run("Meridian Consulting Group")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    for line in [
        "1247 Riverside Boulevard, Suite 300",
        "Portland, OR 97205",
        "Tel: (503) 555-0184",
        "info@meridianconsulting.com",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.name = "Calibri"

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(24)
    date_para.paragraph_format.space_after = Pt(12)
    r = date_para.add_run("March 28, 2026")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # --- Delivery Address (default position - plain paragraphs) ---
    address_lines = [
        "Ms. Elena Vasquez",
        "Director of Operations",
        "Northfield Industries Ltd.",
        "890 Commerce Park Drive",
        "Seattle, WA 98101",
    ]
    for i, line in enumerate(address_lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # --- Salutation ---
    sal = doc.add_paragraph()
    sal.paragraph_format.space_before = Pt(12)
    sal.paragraph_format.space_after = Pt(6)
    r = sal.add_run("Dear Ms. Vasquez,")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # --- Body Paragraphs ---
    body_texts = [
        "Thank you for meeting with our team last Thursday to discuss the proposed "
        "operational efficiency review for Northfield Industries. We were impressed by "
        "your organization's commitment to continuous improvement, and we are confident "
        "that our consulting framework can deliver measurable results within the first quarter.",

        "As discussed, our engagement would encompass three key phases: an initial assessment "
        "of current workflows and bottlenecks (weeks 1\u20133), development of targeted optimization "
        "strategies with stakeholder input (weeks 4\u20136), and supervised implementation with "
        "progress benchmarks (weeks 7\u201312). We anticipate a 15\u201320% reduction in processing "
        "cycle times based on similar engagements in the manufacturing sector.",

        "Enclosed please find our formal proposal outlining the scope of work, projected "
        "timelines, deliverables, and fee structure. The total investment for the twelve-week "
        "engagement is $87,500, which includes all on-site consulting days, interim reports, "
        "and a comprehensive final assessment document.",

        "We would welcome the opportunity to schedule a follow-up call with your leadership "
        "team to address any questions and refine the project scope. Please feel free to "
        "reach me directly at (503) 555-0184 or via email at j.whitfield@meridianconsulting.com.",
    ]

    for text in body_texts:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # --- Closing ---
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    closing.paragraph_format.space_after = Pt(24)
    r = closing.add_run("Sincerely,")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # --- Signature ---
    sig_name = doc.add_paragraph()
    sig_name.paragraph_format.space_before = Pt(0)
    sig_name.paragraph_format.space_after = Pt(0)
    r = sig_name.add_run("James Whitfield")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_before = Pt(0)
    sig_title.paragraph_format.space_after = Pt(0)
    r = sig_title.add_run("Senior Consulting Partner")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    sig_co = doc.add_paragraph()
    sig_co.paragraph_format.space_before = Pt(0)
    sig_co.paragraph_format.space_after = Pt(0)
    r = sig_co.add_run("Meridian Consulting Group")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # --- Enclosure note ---
    enc = doc.add_paragraph()
    enc.paragraph_format.space_before = Pt(12)
    r = enc.add_run("Encl.: Project Proposal \u2013 Northfield Industries Operational Review")
    r.font.size = Pt(10)
    r.italic = True
    r.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
