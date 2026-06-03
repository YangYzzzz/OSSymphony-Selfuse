"""
Initial Setup: Operations Manual with chapter numbering and TOC (without entry numbers)
Task ID: writer_mt_068
Domain: libreoffice_writer

Creates a document with:
- Chapter numbering configured (Heading 1: 1, 2, 3; Heading 2: 1.1, 1.2, etc.)
- A Table of Contents that does NOT show chapter numbers in entries
- Multiple chapters with sub-sections containing realistic content
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_068'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def setup_chapter_numbering(doc):
    """Configure outline numbering for Heading 1 and Heading 2 styles."""
    # Access the numbering part - create if needed
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part._element

    # Create an abstract numbering definition for outline/chapter numbering
    abstract_num_xml = f'''
    <w:abstractNum w:abstractNumId="0" {nsdecls('w')}>
        <w:multiLevelType w:val="multilevel"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:pStyle w:val="Heading1"/>
            <w:lvlText w:val="%1"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="432" w:hanging="432"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:pStyle w:val="Heading2"/>
            <w:lvlText w:val="%1.%2"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="576" w:hanging="576"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="2">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:pStyle w:val="Heading3"/>
            <w:lvlText w:val="%1.%2.%3"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="720" w:hanging="720"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="3">
            <w:start w:val="1"/>
            <w:numFmt w:val="none"/>
            <w:lvlText w:val=""/>
            <w:lvlJc w:val="left"/>
        </w:lvl>
        <w:lvl w:ilvl="4">
            <w:start w:val="1"/>
            <w:numFmt w:val="none"/>
            <w:lvlText w:val=""/>
            <w:lvlJc w:val="left"/>
        </w:lvl>
        <w:lvl w:ilvl="5">
            <w:start w:val="1"/>
            <w:numFmt w:val="none"/>
            <w:lvlText w:val=""/>
            <w:lvlJc w:val="left"/>
        </w:lvl>
        <w:lvl w:ilvl="6">
            <w:start w:val="1"/>
            <w:numFmt w:val="none"/>
            <w:lvlText w:val=""/>
            <w:lvlJc w:val="left"/>
        </w:lvl>
        <w:lvl w:ilvl="7">
            <w:start w:val="1"/>
            <w:numFmt w:val="none"/>
            <w:lvlText w:val=""/>
            <w:lvlJc w:val="left"/>
        </w:lvl>
        <w:lvl w:ilvl="8">
            <w:start w:val="1"/>
            <w:numFmt w:val="none"/>
            <w:lvlText w:val=""/>
            <w:lvlJc w:val="left"/>
        </w:lvl>
    </w:abstractNum>
    '''
    abstract_num = parse_xml(abstract_num_xml)

    # Insert abstractNum before any num elements
    first_num = numbering_elm.find(qn('w:num'))
    if first_num is not None:
        first_num.addprevious(abstract_num)
    else:
        numbering_elm.append(abstract_num)

    # Create a concrete numbering instance referencing the abstract definition
    num_xml = f'''
    <w:num w:numId="1" {nsdecls('w')}>
        <w:abstractNumId w:val="0"/>
    </w:num>
    '''
    num_elm = parse_xml(num_xml)
    numbering_elm.append(num_elm)

    return 1  # numId


def add_heading_with_numbering(doc, text, level, num_id):
    """Add a heading with outline numbering reference."""
    heading = doc.add_heading(text, level=level)
    # Add numbering reference to the paragraph
    pPr = heading._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        heading._element.insert(0, pPr)

    numPr = parse_xml(f'''
        <w:numPr {nsdecls('w')}>
            <w:ilvl w:val="{level - 1}"/>
            <w:numId w:val="{num_id}"/>
        </w:numPr>
    ''')
    pPr.append(numPr)
    return heading


def add_toc_without_entry_numbers(doc):
    """
    Insert a TOC field that does NOT include the chapter number (E#) in entries.
    Uses a standard TOC field code: TOC \\o "1-2" \\h
    This produces entries like: "Introduction.....1" (no chapter number prefix).
    """
    # Create SDT (Structured Document Tag) for TOC
    sdt_xml = f'''
    <w:sdt {nsdecls('w', 'r')}>
        <w:sdtPr>
            <w:docPartObj>
                <w:docPartGallery w:val="Table of Contents"/>
                <w:docPartUnique/>
            </w:docPartObj>
        </w:sdtPr>
        <w:sdtContent>
            <w:p>
                <w:pPr>
                    <w:pStyle w:val="TOCHeading"/>
                    <w:jc w:val="center"/>
                </w:pPr>
                <w:r>
                    <w:rPr>
                        <w:b/>
                        <w:sz w:val="28"/>
                    </w:rPr>
                    <w:t>Table of Contents</w:t>
                </w:r>
            </w:p>
            <w:p>
                <w:r>
                    <w:fldChar w:fldCharType="begin"/>
                </w:r>
                <w:r>
                    <w:instrText xml:space="preserve"> TOC \\o "1-2" \\h </w:instrText>
                </w:r>
                <w:r>
                    <w:fldChar w:fldCharType="separate"/>
                </w:r>
            </w:p>
            <w:p>
                <w:pPr><w:pStyle w:val="TOC1"/></w:pPr>
                <w:r><w:t>Introduction</w:t></w:r>
                <w:r><w:tab/></w:r>
                <w:r><w:fldChar w:fldCharType="begin"/></w:r>
                <w:r><w:instrText xml:space="preserve"> PAGEREF _Toc0001 \\h </w:instrText></w:r>
                <w:r><w:fldChar w:fldCharType="separate"/></w:r>
                <w:r><w:t>1</w:t></w:r>
                <w:r><w:fldChar w:fldCharType="end"/></w:r>
            </w:p>
            <w:p>
                <w:pPr><w:pStyle w:val="TOC2"/></w:pPr>
                <w:r><w:t>Scope</w:t></w:r>
                <w:r><w:tab/></w:r>
                <w:r><w:fldChar w:fldCharType="begin"/></w:r>
                <w:r><w:instrText xml:space="preserve"> PAGEREF _Toc0002 \\h </w:instrText></w:r>
                <w:r><w:fldChar w:fldCharType="separate"/></w:r>
                <w:r><w:t>2</w:t></w:r>
                <w:r><w:fldChar w:fldCharType="end"/></w:r>
            </w:p>
            <w:p>
                <w:pPr><w:pStyle w:val="TOC2"/></w:pPr>
                <w:r><w:t>Purpose</w:t></w:r>
                <w:r><w:tab/></w:r>
                <w:r><w:fldChar w:fldCharType="begin"/></w:r>
                <w:r><w:instrText xml:space="preserve"> PAGEREF _Toc0003 \\h </w:instrText></w:r>
                <w:r><w:fldChar w:fldCharType="separate"/></w:r>
                <w:r><w:t>3</w:t></w:r>
                <w:r><w:fldChar w:fldCharType="end"/></w:r>
            </w:p>
            <w:p>
                <w:pPr><w:pStyle w:val="TOC1"/></w:pPr>
                <w:r><w:t>Procedures</w:t></w:r>
                <w:r><w:tab/></w:r>
                <w:r><w:fldChar w:fldCharType="begin"/></w:r>
                <w:r><w:instrText xml:space="preserve"> PAGEREF _Toc0004 \\h </w:instrText></w:r>
                <w:r><w:fldChar w:fldCharType="separate"/></w:r>
                <w:r><w:t>5</w:t></w:r>
                <w:r><w:fldChar w:fldCharType="end"/></w:r>
            </w:p>
            <w:p>
                <w:pPr><w:pStyle w:val="TOC2"/></w:pPr>
                <w:r><w:t>Standard Operating Procedures</w:t></w:r>
                <w:r><w:tab/></w:r>
                <w:r><w:fldChar w:fldCharType="begin"/></w:r>
                <w:r><w:instrText xml:space="preserve"> PAGEREF _Toc0005 \\h </w:instrText></w:r>
                <w:r><w:fldChar w:fldCharType="separate"/></w:r>
                <w:r><w:t>5</w:t></w:r>
                <w:r><w:fldChar w:fldCharType="end"/></w:r>
            </w:p>
            <w:p>
                <w:pPr><w:pStyle w:val="TOC2"/></w:pPr>
                <w:r><w:t>Emergency Protocols</w:t></w:r>
                <w:r><w:tab/></w:r>
                <w:r><w:fldChar w:fldCharType="begin"/></w:r>
                <w:r><w:instrText xml:space="preserve"> PAGEREF _Toc0006 \\h </w:instrText></w:r>
                <w:r><w:fldChar w:fldCharType="separate"/></w:r>
                <w:r><w:t>7</w:t></w:r>
                <w:r><w:fldChar w:fldCharType="end"/></w:r>
            </w:p>
            <w:p>
                <w:pPr><w:pStyle w:val="TOC1"/></w:pPr>
                <w:r><w:t>Safety Guidelines</w:t></w:r>
                <w:r><w:tab/></w:r>
                <w:r><w:fldChar w:fldCharType="begin"/></w:r>
                <w:r><w:instrText xml:space="preserve"> PAGEREF _Toc0007 \\h </w:instrText></w:r>
                <w:r><w:fldChar w:fldCharType="separate"/></w:r>
                <w:r><w:t>9</w:t></w:r>
                <w:r><w:fldChar w:fldCharType="end"/></w:r>
            </w:p>
            <w:p>
                <w:pPr><w:pStyle w:val="TOC2"/></w:pPr>
                <w:r><w:t>Workplace Safety</w:t></w:r>
                <w:r><w:tab/></w:r>
                <w:r><w:fldChar w:fldCharType="begin"/></w:r>
                <w:r><w:instrText xml:space="preserve"> PAGEREF _Toc0008 \\h </w:instrText></w:r>
                <w:r><w:fldChar w:fldCharType="separate"/></w:r>
                <w:r><w:t>9</w:t></w:r>
                <w:r><w:fldChar w:fldCharType="end"/></w:r>
            </w:p>
            <w:p>
                <w:pPr><w:pStyle w:val="TOC2"/></w:pPr>
                <w:r><w:t>Equipment Handling</w:t></w:r>
                <w:r><w:tab/></w:r>
                <w:r><w:fldChar w:fldCharType="begin"/></w:r>
                <w:r><w:instrText xml:space="preserve"> PAGEREF _Toc0009 \\h </w:instrText></w:r>
                <w:r><w:fldChar w:fldCharType="separate"/></w:r>
                <w:r><w:t>10</w:t></w:r>
                <w:r><w:fldChar w:fldCharType="end"/></w:r>
            </w:p>
            <w:p>
                <w:r>
                    <w:fldChar w:fldCharType="end"/>
                </w:r>
            </w:p>
        </w:sdtContent>
    </w:sdt>
    '''
    sdt = parse_xml(sdt_xml)
    return sdt


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set up chapter numbering
    num_id = setup_chapter_numbering(doc)

    # ---- Title Page ----
    title = doc.add_heading('Operations Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Industrial Solutions, Inc.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Effective Date: March 15, 2025')
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph()  # spacing

    # ---- Insert TOC (without chapter numbers in entries) ----
    toc_sdt = add_toc_without_entry_numbers(doc)
    # Insert TOC after the spacing paragraph
    doc.element.body.append(toc_sdt)

    # Page break after TOC
    doc.add_page_break()

    # ---- Chapter 1: Introduction ----
    h1 = add_heading_with_numbering(doc, 'Introduction', 1, num_id)

    doc.add_paragraph(
        'This Operations Manual provides comprehensive guidelines for all departments '
        'within Meridian Industrial Solutions, Inc. It serves as the primary reference '
        'document for operational procedures, safety protocols, and quality standards '
        'that govern our daily business activities.'
    )
    doc.add_paragraph(
        'All employees are expected to familiarize themselves with the contents of this '
        'manual and adhere to the procedures outlined herein. Failure to comply with '
        'established protocols may result in disciplinary action as described in the '
        'Employee Handbook (Section 7.3).'
    )

    # ---- Section 1.1: Scope ----
    h1_1 = add_heading_with_numbering(doc, 'Scope', 2, num_id)

    doc.add_paragraph(
        'This manual applies to all full-time and part-time employees across the following '
        'operational divisions:'
    )
    doc.add_paragraph('Manufacturing and Production', style='List Bullet')
    doc.add_paragraph('Quality Assurance and Control', style='List Bullet')
    doc.add_paragraph('Warehouse and Logistics', style='List Bullet')
    doc.add_paragraph('Facility Maintenance', style='List Bullet')
    doc.add_paragraph('Health, Safety, and Environment (HSE)', style='List Bullet')
    doc.add_paragraph(
        'Third-party contractors and temporary workers must also comply with relevant '
        'sections of this manual, particularly those pertaining to safety and access control.'
    )

    # ---- Section 1.2: Purpose ----
    h1_2 = add_heading_with_numbering(doc, 'Purpose', 2, num_id)

    doc.add_paragraph(
        'The purpose of this document is to establish a unified set of operational standards '
        'that ensure consistency, efficiency, and regulatory compliance across all Meridian '
        'facilities. Key objectives include:'
    )
    doc.add_paragraph('Standardizing operating procedures across 14 regional facilities', style='List Number')
    doc.add_paragraph('Reducing workplace incidents by 25% year-over-year', style='List Number')
    doc.add_paragraph('Achieving ISO 9001:2015 compliance by Q4 2025', style='List Number')
    doc.add_paragraph('Improving cross-departmental communication efficiency', style='List Number')

    doc.add_page_break()

    # ---- Chapter 2: Procedures ----
    h2 = add_heading_with_numbering(doc, 'Procedures', 1, num_id)

    doc.add_paragraph(
        'All operational procedures must be documented, reviewed annually, and approved '
        'by the relevant department head before implementation. The procedures described '
        'in this chapter represent the minimum standards required for compliance.'
    )

    # ---- Section 2.1: Standard Operating Procedures ----
    h2_1 = add_heading_with_numbering(doc, 'Standard Operating Procedures', 2, num_id)

    doc.add_paragraph(
        'Standard Operating Procedures (SOPs) are detailed, written instructions designed '
        'to achieve uniformity in the performance of specific functions. Each SOP must include:'
    )
    doc.add_paragraph('A clear objective statement and scope definition', style='List Bullet')
    doc.add_paragraph('Step-by-step instructions with decision points', style='List Bullet')
    doc.add_paragraph('Required equipment, materials, and personal protective equipment (PPE)', style='List Bullet')
    doc.add_paragraph('Quality checkpoints and acceptance criteria', style='List Bullet')
    doc.add_paragraph('Revision history and approval signatures', style='List Bullet')

    doc.add_paragraph(
        'SOPs are maintained in the central document management system (DMS) and must be '
        'reviewed at minimum every 12 months or when a significant process change occurs. '
        'The Quality Assurance team is responsible for auditing SOP compliance quarterly.'
    )

    # Add a table for SOP tracking
    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['SOP ID', 'Description', 'Last Review', 'Next Review']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    sop_data = [
        ['SOP-2025-001', 'Assembly Line Startup Sequence', '2025-01-10', '2026-01-10'],
        ['SOP-2025-002', 'Chemical Storage and Handling', '2025-02-15', '2026-02-15'],
        ['SOP-2025-003', 'Quality Inspection Protocol', '2025-03-01', '2026-03-01'],
        ['SOP-2025-004', 'Forklift Operation Standards', '2024-12-20', '2025-12-20'],
    ]
    for r, row_data in enumerate(sop_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_page_break()

    # ---- Section 2.2: Emergency Protocols ----
    h2_2 = add_heading_with_numbering(doc, 'Emergency Protocols', 2, num_id)

    doc.add_paragraph(
        'Emergency protocols are activated when an incident occurs that poses immediate '
        'risk to personnel, property, or the environment. All employees must complete '
        'annual emergency response training and participate in quarterly evacuation drills.'
    )
    doc.add_paragraph(
        'The Emergency Response Team (ERT) consists of trained volunteers from each '
        'department, coordinated by the HSE Director. Response times must meet the '
        'following benchmarks:'
    )
    doc.add_paragraph('Fire alarm response: evacuation within 3 minutes', style='List Bullet')
    doc.add_paragraph('Chemical spill containment: initial response within 5 minutes', style='List Bullet')
    doc.add_paragraph('Medical emergency: first aid within 2 minutes, EMT arrival within 8 minutes', style='List Bullet')
    doc.add_paragraph('Severe weather shelter-in-place: compliance within 4 minutes', style='List Bullet')

    doc.add_paragraph(
        'All emergency incidents must be reported within 24 hours using Form ER-100, '
        'available on the company intranet. The HSE Director will conduct a root cause '
        'analysis within 5 business days of any Level 2 or higher incident.'
    )

    doc.add_page_break()

    # ---- Chapter 3: Safety Guidelines ----
    h3 = add_heading_with_numbering(doc, 'Safety Guidelines', 1, num_id)

    doc.add_paragraph(
        'Meridian Industrial Solutions is committed to maintaining a zero-harm workplace. '
        'Safety is not merely a priority but a core value embedded in every operational '
        'decision. These guidelines establish the minimum safety standards for all facilities.'
    )

    # ---- Section 3.1: Workplace Safety ----
    h3_1 = add_heading_with_numbering(doc, 'Workplace Safety', 2, num_id)

    doc.add_paragraph(
        'General workplace safety encompasses the physical environment, ergonomic standards, '
        'and behavioral expectations for all personnel. Key requirements include:'
    )
    doc.add_paragraph('Mandatory PPE in all production areas (hard hat, safety glasses, steel-toed boots)', style='List Bullet')
    doc.add_paragraph('Clear and unobstructed emergency exits at all times', style='List Bullet')
    doc.add_paragraph('Proper labeling of all hazardous materials per OSHA GHS standards', style='List Bullet')
    doc.add_paragraph('Maximum noise exposure of 85 dB TWA without hearing protection', style='List Bullet')
    doc.add_paragraph(
        'Monthly safety inspections are conducted by department safety liaisons using '
        'the standardized checklist (Form SI-200). Results are reported to the HSE '
        'committee and tracked in the safety performance dashboard.'
    )

    # ---- Section 3.2: Equipment Handling ----
    h3_2 = add_heading_with_numbering(doc, 'Equipment Handling', 2, num_id)

    doc.add_paragraph(
        'All equipment operators must complete the appropriate training and certification '
        'program before operating any machinery. Equipment-specific requirements include:'
    )
    doc.add_paragraph('Forklift certification (renewed every 3 years) per OSHA 1910.178', style='List Bullet')
    doc.add_paragraph('Overhead crane operation license with annual competency assessment', style='List Bullet')
    doc.add_paragraph('Lockout/Tagout (LOTO) training for all maintenance personnel', style='List Bullet')
    doc.add_paragraph('Confined space entry permit and rescue team certification', style='List Bullet')

    doc.add_paragraph(
        'Equipment maintenance logs must be updated after every service event. '
        'Preventive maintenance schedules are managed through the CMMS (Computerized '
        'Maintenance Management System) and tracked by the Maintenance Supervisor. '
        'Any equipment found to be in unsafe condition must be immediately tagged out '
        'and reported via Form EQ-300.'
    )

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
