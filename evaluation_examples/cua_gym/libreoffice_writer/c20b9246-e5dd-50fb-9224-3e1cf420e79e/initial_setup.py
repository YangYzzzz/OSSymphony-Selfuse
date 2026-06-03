"""
Initial Setup: Leave policy document with introductory text (no table yet)
Task ID: writer_hr_030
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_030'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Leave Policy 2026', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies Inc.')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'

    doc.add_paragraph()  # spacer

    # --- Introduction ---
    intro1 = doc.add_paragraph()
    run = intro1.add_run(
        'Meridian Technologies is committed to supporting the well-being and work-life balance '
        'of all employees. Our comprehensive leave program is designed to provide adequate time '
        'off for rest, personal matters, family needs, and unforeseen circumstances. This policy '
        'outlines the various types of leave available and the conditions under which they may be used.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    intro2 = doc.add_paragraph()
    run = intro2.add_run(
        'All leave requests must be submitted through the HR portal at least two weeks in advance '
        'for planned absences. Emergency or unplanned leave should be reported to your direct '
        'supervisor and the HR department within 24 hours. Approval of leave requests is subject '
        'to departmental staffing requirements and manager discretion.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    intro3 = doc.add_paragraph()
    run = intro3.add_run(
        'The following table summarizes our current leave categories, annual entitlements, '
        'eligibility criteria, and carryover provisions. Employees are encouraged to familiarize '
        'themselves with each category and plan their time off accordingly.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Footer note ---
    doc.add_paragraph()  # spacer
    footer_para = doc.add_paragraph()
    run = footer_para.add_run(
        'For questions about this policy, please contact the Human Resources department '
        'at hr@meridiantech.com or extension 4200.'
    )
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.italic = True

    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
