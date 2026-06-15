"""
Initial Setup: HR Newsletter in single-column layout
Task ID: writer_hr_034
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_034'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Newsletter Title / Header ---
    title = doc.add_heading('', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('Meridian Solutions HR Newsletter')
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_sub = subtitle.add_run('April 2025 Edition  |  Volume 12, Issue 4')
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run_sub.italic = True

    # Divider line
    divider = doc.add_paragraph()
    divider.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    div_run = divider.add_run('\u2500' * 60)
    div_run.font.size = Pt(8)
    div_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # --- Article 1 ---
    h1 = doc.add_heading('New Employee Wellness Program Launches May 1st', level=2)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p1a = doc.add_paragraph()
    r1a = p1a.add_run(
        'We are excited to announce the launch of our comprehensive Employee Wellness Program, '
        'developed in partnership with HealthBridge Consulting. Starting May 1st, all full-time '
        'employees will have access to a range of wellness resources designed to support physical, '
        'mental, and financial well-being.'
    )
    r1a.font.size = Pt(11)

    p1b = doc.add_paragraph()
    r1b = p1b.add_run(
        'The program includes on-site yoga sessions every Tuesday and Thursday from 12:00 PM to '
        '12:45 PM in Conference Room B, a 24/7 mental health support hotline staffed by licensed '
        'counselors, monthly financial planning workshops led by Greenfield Advisors, and subsidized '
        'gym memberships at FitLife Centers across all office locations. Enrollment is automatic for '
        'eligible employees, but you can customize your preferences through the HR portal.'
    )
    r1b.font.size = Pt(11)

    p1c = doc.add_paragraph()
    r1c = p1c.add_run(
        'Director of People Operations, Vanessa Liu, noted: "Our team surveyed over 800 employees '
        'last quarter, and the results were clear \u2014 wellness support is a top priority. We listened, '
        'and this program is the direct result of that feedback." For questions, contact the Wellness '
        'Team at wellness@meridiansolutions.com.'
    )
    r1c.font.size = Pt(11)

    # --- Article 2 ---
    h2 = doc.add_heading('Q1 Performance Review Highlights', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p2a = doc.add_paragraph()
    r2a = p2a.add_run(
        'The first quarter performance review cycle has concluded, and we are proud to share some '
        'outstanding results. Across all departments, 87% of employees met or exceeded their '
        'performance targets, representing a 4% improvement over Q1 last year.'
    )
    r2a.font.size = Pt(11)

    p2b = doc.add_paragraph()
    r2b = p2b.add_run(
        'Notable achievements include the Engineering team delivering Project Aurora three weeks '
        'ahead of schedule under the leadership of Senior Manager Rajesh Patel, the Sales division '
        'surpassing their quarterly revenue target by 12% thanks to the efforts of Regional Director '
        'Camille Dubois, and the Customer Success team achieving a record-setting Net Promoter Score '
        'of 78, up from 71 in Q4. Department heads will schedule individual feedback sessions '
        'throughout April.'
    )
    r2b.font.size = Pt(11)

    p2c = doc.add_paragraph()
    r2c = p2c.add_run(
        'Chief Human Resources Officer, Marcus Thompson, commented: "These numbers reflect the '
        'dedication and talent across every level of our organization. I encourage all managers to '
        'recognize their teams and continue fostering a culture of excellence." Detailed department '
        'reports are available on the internal HR dashboard.'
    )
    r2c.font.size = Pt(11)

    # --- Article 3 ---
    h3 = doc.add_heading('Upcoming Company Events and Important Dates', level=2)
    for run in h3.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p3a = doc.add_paragraph()
    r3a = p3a.add_run(
        'Mark your calendars for an exciting lineup of company events this spring. On May 10th, '
        'we will host our annual Meridian Family Day at Riverside Park, featuring games, food trucks, '
        'and a live performance by local band The Copper Keys. RSVP through the Events page on the '
        'intranet by April 25th.'
    )
    r3a.font.size = Pt(11)

    p3b = doc.add_paragraph()
    r3b = p3b.add_run(
        'The Professional Development Summit is scheduled for May 22nd through May 23rd at the '
        'downtown Marriott Conference Center. This year\'s theme is "Leading Through Change," with '
        'keynote speakers including Dr. Anika Sharma from Stanford Graduate School of Business and '
        'tech industry leader Devon Park, CEO of NexGen Dynamics. Registration opens April 15th on '
        'the Learning & Development portal.'
    )
    r3b.font.size = Pt(11)

    p3c = doc.add_paragraph()
    r3c = p3c.add_run(
        'Other important dates to remember: April 18th is the deadline for submitting summer '
        'vacation requests, May 5th is the open enrollment deadline for updated health insurance '
        'plans, and May 15th is the last day to nominate colleagues for the quarterly Meridian Star '
        'Award. Please reach out to events@meridiansolutions.com with any questions.'
    )
    r3c.font.size = Pt(11)

    # --- Footer divider and footer text ---
    divider2 = doc.add_paragraph()
    divider2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    div2_run = divider2.add_run('\u2500' * 60)
    div2_run.font.size = Pt(8)
    div2_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fr = footer_p.add_run(
        'Meridian Solutions  |  1200 Innovation Drive, Suite 400, San Francisco, CA 94105  |  '
        'hr@meridiansolutions.com'
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    footer_p2 = doc.add_paragraph()
    footer_p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fr2 = footer_p2.add_run('Confidential \u2014 For Internal Distribution Only')
    fr2.font.size = Pt(8)
    fr2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    fr2.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
