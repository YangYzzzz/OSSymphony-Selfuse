"""
Initial Setup: Legal document with 10 tracked changes for selective review
Task ID: writer_lec_070
Domain: libreoffice_writer

Creates a legal services agreement with 10 tracked changes from a reviewer.
Changes 3, 5, and 8 modify liability limitation clauses.
"""

import os
import shlex
import subprocess
import time
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_070'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

AUTHOR = "Jennifer Walsh"
DATE = "2025-11-10T14:30:00Z"

# Word namespace
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_run_element(text, bold=False, italic=False, font_name="Times New Roman", font_size_pt=12):
    """Create a w:r element with optional formatting."""
    r = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{text}</w:t></w:r>')
    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
    if bold:
        rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
    if italic:
        rPr.append(parse_xml(f'<w:i {nsdecls("w")}/>'))
    if font_name:
        rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}"/>'))
    if font_size_pt:
        half_pt = font_size_pt * 2
        rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{half_pt}"/>'))
    r.insert(0, rPr)
    return r


def add_tracked_insertion(para_element, rev_id, text, bold=False, italic=False,
                          font_name="Times New Roman", font_size_pt=12, append=True):
    """Add a tracked insertion (w:ins) to a paragraph element."""
    ins = parse_xml(
        f'<w:ins {nsdecls("w")} w:id="{rev_id}" w:author="{AUTHOR}" w:date="{DATE}"/>'
    )
    r = make_run_element(text, bold=bold, italic=italic, font_name=font_name, font_size_pt=font_size_pt)
    ins.append(r)
    if append:
        para_element.append(ins)
    return ins


def add_tracked_deletion(para_element, rev_id, text, bold=False, italic=False,
                         font_name="Times New Roman", font_size_pt=12, append=True):
    """Add a tracked deletion (w:del) to a paragraph element."""
    del_elem = parse_xml(
        f'<w:del {nsdecls("w")} w:id="{rev_id}" w:author="{AUTHOR}" w:date="{DATE}"/>'
    )
    r = parse_xml(f'<w:r {nsdecls("w")}><w:delText xml:space="preserve">{text}</w:delText></w:r>')
    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
    if bold:
        rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
    if italic:
        rPr.append(parse_xml(f'<w:i {nsdecls("w")}/>'))
    if font_name:
        rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}"/>'))
    if font_size_pt:
        half_pt = font_size_pt * 2
        rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{half_pt}"/>'))
    r.insert(0, rPr)
    del_elem.append(r)
    if append:
        para_element.append(del_elem)
    return del_elem


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ============================================================
    # TITLE
    # ============================================================
    title = doc.add_heading("PROFESSIONAL SERVICES AGREEMENT", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")  # spacer

    # ============================================================
    # PREAMBLE
    # ============================================================
    preamble = doc.add_paragraph(
        'This Professional Services Agreement ("Agreement") is entered into as of '
        'November 1, 2025, by and between Meridian Consulting Group, LLC ("Service Provider") '
        'and Blackstone Enterprises, Inc. ("Client"), collectively referred to as the "Parties."'
    )
    preamble.paragraph_format.space_after = Pt(6)

    # ============================================================
    # SECTION 1 — SCOPE OF SERVICES
    # ============================================================
    doc.add_heading("1. SCOPE OF SERVICES", level=1)

    # Paragraph with Change 1: "quarterly" → "monthly" (non-liability, accept)
    p1 = doc.add_paragraph()
    run1a = p1.add_run(
        "The Service Provider shall deliver comprehensive strategic consulting services, "
        "including market analysis, operational assessments, and financial advisory. "
        "The Service Provider shall submit detailed progress reports on a "
    )
    # Change 1: delete "quarterly" insert "monthly"
    p1_elem = p1._element
    add_tracked_deletion(p1_elem, 1, "quarterly")
    add_tracked_insertion(p1_elem, 2, "monthly")
    run1b = p1.add_run(
        " basis throughout the engagement period."
    )
    p1.paragraph_format.space_after = Pt(6)

    # ============================================================
    # SECTION 2 — COMPENSATION
    # ============================================================
    doc.add_heading("2. COMPENSATION AND PAYMENT TERMS", level=1)

    # Paragraph with Change 2: "Net 30" → "Net 45" (non-liability, accept)
    p2 = doc.add_paragraph()
    run2a = p2.add_run(
        "Client agrees to compensate the Service Provider at a rate of $275.00 per hour "
        "for all authorized services rendered. Payment shall be due within "
    )
    p2_elem = p2._element
    add_tracked_deletion(p2_elem, 3, "Net 30")
    add_tracked_insertion(p2_elem, 4, "Net 45")
    run2b = p2.add_run(
        " days of invoice receipt. Late payments shall accrue interest at a rate of "
        "1.5% per month."
    )
    p2.paragraph_format.space_after = Pt(6)

    # ============================================================
    # SECTION 3 — CONFIDENTIALITY
    # ============================================================
    doc.add_heading("3. CONFIDENTIALITY", level=1)

    # Paragraph with Change 3 (LIABILITY-RELATED — reject):
    # Tries to change "two (2) years" → "one (1) year" for confidentiality obligations
    p3 = doc.add_paragraph()
    run3a = p3.add_run(
        "Both Parties agree to maintain strict confidentiality regarding all proprietary "
        "information, trade secrets, and business strategies disclosed during the engagement. "
        "Confidentiality obligations shall survive termination of this Agreement for a period of "
    )
    p3_elem = p3._element
    add_tracked_deletion(p3_elem, 5, "two (2) years")
    add_tracked_insertion(p3_elem, 6, "one (1) year")
    run3c = p3.add_run(
        " from the date of termination. Breach of this provision shall entitle the "
        "non-breaching Party to seek injunctive relief and monetary damages."
    )
    p3.paragraph_format.space_after = Pt(6)

    # ============================================================
    # SECTION 4 — INTELLECTUAL PROPERTY
    # ============================================================
    doc.add_heading("4. INTELLECTUAL PROPERTY", level=1)

    # Paragraph with Change 4: "upon full payment" → "upon delivery" (non-liability, accept)
    p4 = doc.add_paragraph()
    run4a = p4.add_run(
        "All deliverables, reports, and work products created by the Service Provider "
        "in the course of this engagement shall become the exclusive property of the Client "
    )
    p4_elem = p4._element
    add_tracked_deletion(p4_elem, 7, "upon full payment")
    add_tracked_insertion(p4_elem, 8, "upon delivery")
    run4b = p4.add_run(
        ". The Service Provider retains no rights to use, reproduce, or distribute "
        "such materials without prior written consent from the Client."
    )
    p4.paragraph_format.space_after = Pt(6)

    # ============================================================
    # SECTION 5 — LIABILITY LIMITATION
    # ============================================================
    doc.add_heading("5. LIMITATION OF LIABILITY", level=1)

    # Paragraph with Change 5 (LIABILITY-RELATED — reject):
    # Tries to change liability cap from "total fees paid" to "$50,000"
    p5 = doc.add_paragraph()
    run5a = p5.add_run(
        "In no event shall the Service Provider be liable for any indirect, incidental, "
        "consequential, or punitive damages arising out of this Agreement. The total aggregate "
        "liability of the Service Provider under this Agreement shall not exceed "
    )
    p5_elem = p5._element
    add_tracked_deletion(p5_elem, 9, "the total fees paid by Client under this Agreement")
    add_tracked_insertion(p5_elem, 10, "fifty thousand dollars ($50,000)")
    run5b = p5.add_run(
        ", regardless of the cause of action or theory of liability."
    )
    p5.paragraph_format.space_after = Pt(6)

    # ============================================================
    # SECTION 6 — TERM AND TERMINATION
    # ============================================================
    doc.add_heading("6. TERM AND TERMINATION", level=1)

    # Paragraph with Change 6: "thirty (30) days" → "sixty (60) days" notice (non-liability, accept)
    p6 = doc.add_paragraph()
    run6a = p6.add_run(
        "This Agreement shall commence on the Effective Date and continue for an initial "
        "term of twelve (12) months. Either Party may terminate this Agreement with "
    )
    p6_elem = p6._element
    add_tracked_deletion(p6_elem, 11, "thirty (30) days")
    add_tracked_insertion(p6_elem, 12, "sixty (60) days")
    run6b = p6.add_run(
        " prior written notice. Upon termination, Client shall pay for all services "
        "rendered through the effective date of termination."
    )
    p6.paragraph_format.space_after = Pt(6)

    # ============================================================
    # SECTION 7 — INDEMNIFICATION
    # ============================================================
    doc.add_heading("7. INDEMNIFICATION", level=1)

    # Paragraph with Change 7: "negligence or willful misconduct" → "negligence, willful misconduct, or breach of this Agreement"
    # (non-liability, accept — broadens indemnification trigger)
    p7 = doc.add_paragraph()
    run7a = p7.add_run(
        "Each Party agrees to indemnify and hold harmless the other Party from and against "
        "all claims, damages, losses, and expenses (including reasonable attorneys' fees) arising "
        "from the indemnifying Party's "
    )
    p7_elem = p7._element
    add_tracked_deletion(p7_elem, 13, "negligence or willful misconduct")
    add_tracked_insertion(p7_elem, 14, "negligence, willful misconduct, or breach of this Agreement")
    run7b = p7.add_run(
        " in connection with the performance of obligations under this Agreement."
    )
    p7.paragraph_format.space_after = Pt(6)

    # ============================================================
    # SECTION 8 — INSURANCE AND LIABILITY COVERAGE
    # ============================================================
    doc.add_heading("8. INSURANCE AND LIABILITY COVERAGE", level=1)

    # Paragraph with Change 8 (LIABILITY-RELATED — reject):
    # Tries to reduce insurance requirement from "$2,000,000" to "$500,000"
    p8 = doc.add_paragraph()
    run8a = p8.add_run(
        "The Service Provider shall maintain professional liability insurance with coverage "
        "of not less than "
    )
    p8_elem = p8._element
    add_tracked_deletion(p8_elem, 15, "two million dollars ($2,000,000)")
    add_tracked_insertion(p8_elem, 16, "five hundred thousand dollars ($500,000)")
    run8b = p8.add_run(
        " per occurrence and in the aggregate for the duration of this Agreement "
        "and for a period of two (2) years following its termination."
    )
    p8.paragraph_format.space_after = Pt(6)

    # ============================================================
    # SECTION 9 — DISPUTE RESOLUTION
    # ============================================================
    doc.add_heading("9. DISPUTE RESOLUTION", level=1)

    # Paragraph with Change 9: "litigation in state court" → "binding arbitration"
    # (non-liability, accept)
    p9 = doc.add_paragraph()
    run9a = p9.add_run(
        "Any disputes arising under this Agreement shall be resolved through "
    )
    p9_elem = p9._element
    add_tracked_deletion(p9_elem, 17, "litigation in the state courts of New York")
    add_tracked_insertion(p9_elem, 18, "binding arbitration administered by the American Arbitration Association in New York")
    run9b = p9.add_run(
        ". The prevailing Party shall be entitled to recover reasonable attorneys' fees and costs."
    )
    p9.paragraph_format.space_after = Pt(6)

    # ============================================================
    # SECTION 10 — GOVERNING LAW
    # ============================================================
    doc.add_heading("10. GOVERNING LAW AND MISCELLANEOUS", level=1)

    # Paragraph with Change 10: "State of New York" → "State of Delaware" (non-liability, accept)
    p10 = doc.add_paragraph()
    run10a = p10.add_run(
        "This Agreement shall be governed by and construed in accordance with the laws of the "
    )
    p10_elem = p10._element
    add_tracked_deletion(p10_elem, 19, "State of New York")
    add_tracked_insertion(p10_elem, 20, "State of Delaware")
    run10b = p10.add_run(
        ", without regard to its conflict of laws principles. This Agreement constitutes "
        "the entire agreement between the Parties and supersedes all prior negotiations, "
        "representations, or agreements relating to its subject matter."
    )
    p10.paragraph_format.space_after = Pt(6)

    # ============================================================
    # SIGNATURE BLOCK
    # ============================================================
    doc.add_paragraph("")
    doc.add_paragraph("_" * 50)
    sig1 = doc.add_paragraph("For Meridian Consulting Group, LLC")
    sig1.paragraph_format.space_after = Pt(2)
    doc.add_paragraph("Name: Robert A. Thornton")
    doc.add_paragraph("Title: Managing Partner")
    doc.add_paragraph("Date: _______________")

    doc.add_paragraph("")
    doc.add_paragraph("_" * 50)
    sig2 = doc.add_paragraph("For Blackstone Enterprises, Inc.")
    sig2.paragraph_format.space_after = Pt(2)
    doc.add_paragraph("Name: Victoria S. Nakamura")
    doc.add_paragraph("Title: General Counsel")
    doc.add_paragraph("Date: _______________")

    # Enable tracked changes display in document settings
    # Add revision tracking settings to document
    settings_elem = doc.settings.element
    # Turn on revision tracking
    rev_view = parse_xml(f'<w:revisionView {nsdecls("w")} w:markup="1"/>')
    settings_elem.append(rev_view)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
