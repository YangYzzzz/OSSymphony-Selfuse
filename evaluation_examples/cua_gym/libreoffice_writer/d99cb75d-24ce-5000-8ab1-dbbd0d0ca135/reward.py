"""
Reward Script: Review tracked changes in legal document - reject changes 3, 5, 8; accept all others
Task ID: writer_lec_070
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): All tracked changes resolved (no remaining revisions)
  Component 2 (0.15): Change 3 rejected - "two (2) years" preserved in confidentiality clause
  Component 3 (0.15): Change 5 rejected - "total fees paid by Client" preserved in liability cap
  Component 4 (0.15): Change 8 rejected - "two million dollars ($2,000,000)" preserved in insurance
  Component 5 (0.05): Change 1 accepted - "monthly" in progress reports
  Component 6 (0.05): Change 2 accepted - "Net 45" in payment terms
  Component 7 (0.05): Change 4 accepted - "upon delivery" in IP clause
  Component 8 (0.05): Change 6 accepted - "sixty (60) days" in termination notice
  Component 9 (0.05): Change 7 accepted - "breach of this Agreement" in indemnification
  Component 10 (0.05): Change 9 accepted - "binding arbitration" in dispute resolution
  Component 11 (0.05): Change 10 accepted - "State of Delaware" in governing law
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_070'


def persist_app_state(domain):
    """Save any unsaved edits in LibreOffice Writer."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_writer", "libreoffice_calc", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document
    from lxml import etree

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body = doc.element.body

    # Component 1: All tracked changes resolved (0.20 points)
    # Initial doc has 10 insertions + 10 deletions = 20 revision marks
    # Golden doc should have 0
    try:
        inserts = body.findall('.//w:ins', ns)
        deletes = body.findall('.//w:del', ns)
        num_revisions = len(inserts) + len(deletes)
        if num_revisions == 0:
            print(f"PASS: Component 1 - All tracked changes resolved (0 revisions remaining) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 - {num_revisions} tracked changes still present (expected 0)")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Helper: get full paragraph text for a given paragraph index
    def get_para_text(idx):
        if idx < len(doc.paragraphs):
            return doc.paragraphs[idx].text
        return ""

    # Component 2: Change 3 REJECTED - confidentiality period stays "two (2) years" (0.15 pts)
    # Paragraph 8 (Section 3. CONFIDENTIALITY) should contain "two (2) years"
    # The proposed change was to shorten to "one (1) year"
    try:
        para8_text = get_para_text(8)
        if "two (2) years" in para8_text and "one (1) year" not in para8_text:
            print(f"PASS: Component 2 - Change 3 rejected: 'two (2) years' preserved in confidentiality (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 - Expected 'two (2) years' in para 8, found: {repr(para8_text[:100])}")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Change 5 REJECTED - liability cap stays "total fees paid by Client" (0.15 pts)
    # Paragraph 12 (Section 5. LIMITATION OF LIABILITY)
    # The proposed change was to "$50,000" cap
    try:
        para12_text = get_para_text(12)
        if "total fees paid by Client" in para12_text and "fifty thousand" not in para12_text.lower():
            print(f"PASS: Component 3 - Change 5 rejected: original liability cap preserved (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 - Expected 'total fees paid by Client' in para 12, found: {repr(para12_text[:100])}")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Change 8 REJECTED - insurance stays "two million dollars ($2,000,000)" (0.15 pts)
    # Paragraph 18 (Section 8. INSURANCE AND LIABILITY COVERAGE)
    # The proposed change was to reduce to $500,000
    try:
        para18_text = get_para_text(18)
        if "two million dollars" in para18_text and "five hundred thousand" not in para18_text.lower():
            print(f"PASS: Component 4 - Change 8 rejected: '$2M insurance' preserved (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 4 - Expected 'two million dollars' in para 18, found: {repr(para18_text[:100])}")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Change 1 ACCEPTED - "monthly" in progress reports (0.05 pts)
    # Paragraph 4 (Section 1. SCOPE OF SERVICES)
    try:
        para4_text = get_para_text(4)
        if "monthly" in para4_text.lower() and "quarterly" not in para4_text.lower():
            print(f"PASS: Component 5 - Change 1 accepted: 'monthly' reports (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 - Expected 'monthly' (not 'quarterly') in para 4, found: {repr(para4_text[:80])}")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: Change 2 ACCEPTED - "Net 45" in payment terms (0.05 pts)
    # Paragraph 6 (Section 2. COMPENSATION AND PAYMENT TERMS)
    try:
        para6_text = get_para_text(6)
        if "Net 45" in para6_text and "Net 30" not in para6_text:
            print(f"PASS: Component 6 - Change 2 accepted: 'Net 45' payment terms (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 - Expected 'Net 45' (not 'Net 30') in para 6, found: {repr(para6_text[:80])}")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    # Component 7: Change 4 ACCEPTED - "upon delivery" in IP clause (0.05 pts)
    # Paragraph 10 (Section 4. INTELLECTUAL PROPERTY)
    try:
        para10_text = get_para_text(10)
        if "upon delivery" in para10_text and "upon full payment" not in para10_text:
            print(f"PASS: Component 7 - Change 4 accepted: 'upon delivery' IP transfer (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 7 - Expected 'upon delivery' (not 'upon full payment') in para 10")
    except Exception as e:
        print(f"ERROR: Component 7 - {e}")

    # Component 8: Change 6 ACCEPTED - "sixty (60) days" termination notice (0.05 pts)
    # Paragraph 14 (Section 6. TERM AND TERMINATION)
    try:
        para14_text = get_para_text(14)
        if "sixty (60) days" in para14_text and "thirty (30) days" not in para14_text:
            print(f"PASS: Component 8 - Change 6 accepted: 'sixty (60) days' notice (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 8 - Expected 'sixty (60) days' (not 'thirty (30)') in para 14")
    except Exception as e:
        print(f"ERROR: Component 8 - {e}")

    # Component 9: Change 7 ACCEPTED - "breach of this Agreement" added to indemnification (0.05 pts)
    # Paragraph 16 (Section 7. INDEMNIFICATION)
    try:
        para16_text = get_para_text(16)
        if "breach of this Agreement" in para16_text:
            print(f"PASS: Component 9 - Change 7 accepted: 'breach of this Agreement' in indemnification (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 9 - Expected 'breach of this Agreement' in para 16")
    except Exception as e:
        print(f"ERROR: Component 9 - {e}")

    # Component 10: Change 9 ACCEPTED - "binding arbitration" in dispute resolution (0.05 pts)
    # Paragraph 20 (Section 9. DISPUTE RESOLUTION)
    try:
        para20_text = get_para_text(20)
        if "binding arbitration" in para20_text and "litigation" not in para20_text.lower():
            print(f"PASS: Component 10 - Change 9 accepted: 'binding arbitration' (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 10 - Expected 'binding arbitration' (not 'litigation') in para 20")
    except Exception as e:
        print(f"ERROR: Component 10 - {e}")

    # Component 11: Change 10 ACCEPTED - "State of Delaware" in governing law (0.05 pts)
    # Paragraph 22 (Section 10. GOVERNING LAW AND MISCELLANEOUS)
    try:
        para22_text = get_para_text(22)
        if "State of Delaware" in para22_text and "State of New York" not in para22_text:
            print(f"PASS: Component 11 - Change 10 accepted: 'State of Delaware' (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 11 - Expected 'State of Delaware' (not 'State of New York') in para 22")
    except Exception as e:
        print(f"ERROR: Component 11 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
