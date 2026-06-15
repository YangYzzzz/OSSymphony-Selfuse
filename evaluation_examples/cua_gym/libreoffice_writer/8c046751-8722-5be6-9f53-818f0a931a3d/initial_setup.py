"""
Initial Setup: Market research report with inline citations but no bibliography
Task ID: writer_biz_077
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_077'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Cloud Computing Adoption in Mid-Size Enterprises: A Market Research Report', level=0)

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=1)
    p = doc.add_paragraph(
        'This report examines the current state of cloud computing adoption among mid-size '
        'enterprises (500-5,000 employees) across North America and Western Europe. Based on '
        'survey data collected from 1,247 technology decision-makers between January and March '
        '2025, our findings reveal significant shifts in infrastructure spending priorities. '
        'Recent studies have confirmed that hybrid cloud strategies now dominate enterprise '
        'planning [1], while security concerns remain the primary barrier to full migration [2].'
    )

    # --- Market Overview ---
    doc.add_heading('Market Overview', level=1)
    doc.add_paragraph(
        'The global cloud computing market reached $623.3 billion in 2024, representing a '
        '21.7% year-over-year increase. Mid-size enterprises accounted for approximately '
        '34% of this spending, up from 28% in 2022. The acceleration is attributed to '
        'post-pandemic digital transformation initiatives and the declining cost of '
        'infrastructure-as-a-service (IaaS) platforms.'
    )
    doc.add_paragraph(
        'Industry analysts project continued growth at a compound annual rate of 18.3% '
        'through 2028, driven primarily by artificial intelligence workloads and edge '
        'computing requirements [3]. The shift toward multi-cloud architectures has '
        'created new opportunities for integration platform vendors and managed service '
        'providers specializing in cross-cloud orchestration.'
    )

    # --- Key Findings ---
    doc.add_heading('Key Findings', level=1)

    doc.add_heading('Adoption Patterns', level=2)
    doc.add_paragraph(
        'Our survey data indicates that 78.4% of respondents have moved at least one '
        'critical workload to a public cloud provider, compared to 61.2% in our 2023 '
        'survey. Notably, 42.1% of organizations now operate in a multi-cloud environment, '
        'utilizing services from two or more major providers simultaneously.'
    )
    doc.add_paragraph(
        'The distribution of cloud spending across major providers shows continued '
        'concentration, with the top three providers (AWS, Microsoft Azure, and Google '
        'Cloud Platform) capturing 67% of enterprise cloud budgets. However, specialized '
        'providers have gained traction in regulated industries, particularly healthcare '
        'and financial services [4].'
    )

    doc.add_heading('Security and Compliance', level=2)
    doc.add_paragraph(
        'Security remains the most cited concern among technology leaders, with 63.8% '
        'of respondents identifying data protection as their top cloud migration challenge. '
        'Compliance requirements, particularly GDPR and SOC 2 certification, continue to '
        'influence vendor selection and architecture decisions.'
    )
    doc.add_paragraph(
        'Organizations that implemented zero-trust security frameworks reported 47% fewer '
        'security incidents compared to those relying on traditional perimeter-based '
        'approaches. This finding aligns with broader industry research on cloud-native '
        'security architectures [5].'
    )

    # --- Cost Analysis ---
    doc.add_heading('Cost Analysis', level=1)
    doc.add_paragraph(
        'Average annual cloud infrastructure spending among surveyed enterprises reached '
        '$2.4 million in 2024, with a median of $1.8 million. Organizations reported an '
        'average cost savings of 23% compared to equivalent on-premises infrastructure, '
        'though 31% of respondents experienced unexpected cost overruns during their first '
        'year of cloud operations.'
    )

    # --- Table: Cloud Spending by Department ---
    doc.add_heading('Departmental Cloud Spending Breakdown', level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['Department', 'Average Annual Spend', '% of Total Budget']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Engineering & Development', '$845,000', '35.2%'],
        ['Data Analytics & BI', '$432,000', '18.0%'],
        ['IT Operations', '$384,000', '16.0%'],
        ['Sales & Marketing', '$312,000', '13.0%'],
        ['Human Resources', '$216,000', '9.0%'],
        ['Finance & Accounting', '$211,000', '8.8%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Recommendations ---
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph(
        'Based on our analysis, we recommend that mid-size enterprises pursuing cloud '
        'adoption focus on the following strategic priorities:'
    )
    doc.add_paragraph('Develop a comprehensive multi-cloud governance framework before expanding provider relationships.', style='List Bullet')
    doc.add_paragraph('Invest in cloud-native security tooling, particularly zero-trust identity management solutions.', style='List Bullet')
    doc.add_paragraph('Establish FinOps practices early to prevent cost overruns and optimize resource utilization.', style='List Bullet')
    doc.add_paragraph('Prioritize workforce upskilling programs to address the persistent cloud skills gap.', style='List Bullet')

    # --- Conclusion ---
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The cloud computing landscape for mid-size enterprises continues to evolve rapidly. '
        'Organizations that adopt strategic, security-first approaches to cloud migration are '
        'positioned to realize significant competitive advantages. As the market matures, the '
        'emphasis will increasingly shift from initial migration to optimization and innovation '
        'on cloud-native platforms.'
    )

    # NO References section - that is the task for the agent to add

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
