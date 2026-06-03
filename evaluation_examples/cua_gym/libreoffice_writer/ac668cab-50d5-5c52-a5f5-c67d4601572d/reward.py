"""
Reward Script: Apply strikethrough formatting to all text in Section 4: Deleted Content
Task ID: osworld_writer_strikethrough_last_para_004
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): At least one run in Section 4 content paragraph has strikethrough
  Component 2 (0.6): ALL text-bearing runs in Section 4 content paragraph have strikethrough
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_strikethrough_last_para_004'

SECTION4_HEADING_TEXT = 'Section 4: Deleted Content'

def find_section4_content_paragraph(doc):
    """
    Find the content paragraph that immediately follows the 'Section 4: Deleted Content' heading.
    Returns the paragraph object or None if not found.
    """
    paragraphs = list(doc.paragraphs)
    for idx, para in enumerate(paragraphs):
        if SECTION4_HEADING_TEXT in para.text:
            # Look for the next non-empty paragraph after the heading
            for next_para in paragraphs[idx + 1:]:
                if next_para.text.strip():
                    return next_para
    return None


def verify_task(file_path):
    """
    Verify task completion: strikethrough applied to all text in Section 4 content paragraph.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate the Section 4 heading — precondition gate
    section4_heading_found = any(
        SECTION4_HEADING_TEXT in para.text for para in doc.paragraphs
    )
    if not section4_heading_found:
        print(f"CRITICAL: Section 4 heading '{SECTION4_HEADING_TEXT}' not found in document.")
        print("REWARD: 0.0")
        return 0.0

    # Find the Section 4 content paragraph
    section4_para = find_section4_content_paragraph(doc)
    if section4_para is None:
        print("CRITICAL: Could not find a content paragraph following Section 4 heading.")
        print("REWARD: 0.0")
        return 0.0

    # Gather all text-bearing runs in the Section 4 content paragraph
    text_runs = [run for run in section4_para.runs if run.text.strip()]

    if not text_runs:
        print("CRITICAL: Section 4 content paragraph has no text-bearing runs.")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Section 4 content paragraph has {len(text_runs)} text-bearing runs.")

    # Component 1: At least one run in Section 4 content paragraph has strikethrough (0.4 points)
    # This checks partial completion — some text has been struck through.
    try:
        runs_with_strike = [run for run in text_runs if run.font.strike is True]
        if runs_with_strike:
            print(f"PASS: Component 1 — {len(runs_with_strike)}/{len(text_runs)} runs have strikethrough (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — 0/{len(text_runs)} runs have strikethrough; none struck through yet")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: ALL text-bearing runs in Section 4 content paragraph have strikethrough (0.6 points)
    # This checks full completion — every sentence/run is struck through.
    try:
        all_strike = all(run.font.strike is True for run in text_runs)
        if all_strike:
            print(f"PASS: Component 2 — ALL {len(text_runs)} runs have strikethrough (0.6 pts)")
            total_score += 0.6
        else:
            non_strike = [run.text[:40] for run in text_runs if run.font.strike is not True]
            print(f"FAIL: Component 2 — {len(non_strike)} run(s) missing strikethrough: {non_strike}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
