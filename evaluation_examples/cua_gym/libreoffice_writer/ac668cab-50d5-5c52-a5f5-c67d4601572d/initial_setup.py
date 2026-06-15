"""
Initial Setup: Policy document with 5 named sections; Section 4 has 3 sentences (no strikethrough)
Task ID: osworld_writer_strikethrough_last_para_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_strikethrough_last_para_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Document title
    title = doc.add_heading('Company IT Usage Policy', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        'This document outlines the acceptable use policies for all information technology '
        'resources within the organization. All employees are expected to read and comply '
        'with the policies described below. Violations may result in disciplinary action.'
    )

    # Section 1
    doc.add_heading('Section 1: Acceptable Use', level=1)
    doc.add_paragraph(
        'Employees are permitted to use company IT resources for business-related tasks during '
        'working hours. Limited personal use is acceptable provided it does not interfere with '
        'productivity or consume excessive bandwidth. All usage must comply with local, state, '
        'and federal laws and regulations.'
    )

    # Section 2
    doc.add_heading('Section 2: Data Security', level=1)
    doc.add_paragraph(
        'All sensitive company data must be stored on approved, encrypted storage systems. '
        'Employees must use strong, unique passwords for all work accounts and enable '
        'multi-factor authentication where available. Unauthorized sharing of confidential '
        'data outside of the organization is strictly prohibited.'
    )

    # Section 3
    doc.add_heading('Section 3: Network Access', level=1)
    doc.add_paragraph(
        'Access to the corporate network is granted based on role and business need. '
        'Employees must not connect unauthorized devices to the corporate network without '
        'prior approval from the IT department. Use of personal VPN services on company '
        'devices is not permitted.'
    )

    # Section 4: Deleted Content — NO strikethrough in initial state
    doc.add_heading('Section 4: Deleted Content', level=1)
    sec4_para = doc.add_paragraph()
    run1 = sec4_para.add_run(
        'The company previously allowed unrestricted access to social media platforms during work hours. '
    )
    run2 = sec4_para.add_run(
        'Employees were permitted to use personal email accounts on company devices for any purpose. '
    )
    run3 = sec4_para.add_run(
        'Remote work arrangements were previously handled without formal documentation or approval processes.'
    )
    # IMPORTANT: No strikethrough applied to any run in Section 4 in initial state

    # Section 5
    doc.add_heading('Section 5: Enforcement and Compliance', level=1)
    doc.add_paragraph(
        'The IT department reserves the right to monitor network traffic and system usage '
        'to ensure compliance with this policy. Any suspected violations will be reported '
        'to Human Resources and management. Employees found in violation of this policy '
        'may face disciplinary action up to and including termination of employment.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
