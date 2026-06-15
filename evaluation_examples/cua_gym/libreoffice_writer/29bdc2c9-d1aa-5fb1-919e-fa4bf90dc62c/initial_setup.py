"""
Initial Setup: Technical glossary document with term-definition pairs, no tabstops
Task ID: osworld_writer_tabstop_008
Domain: libreoffice_writer

Creates technical_glossary.docx with 18 lines of term+definition pairs.
No tabstops and no tab characters — all text left-aligned on each line.
The agent must add left/center tabstops at 0cm and 8cm, and insert
a tab after the first word so terms stay at left and definitions center at 8cm.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_tabstop_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Remove the default blank paragraph
    # (Document() creates one empty paragraph by default)
    # We'll just add our own paragraphs directly.

    # 18 lines of term-definition pairs (realistic technical glossary content)
    # Format: "TERM Definition text" — one-word term followed by its definition
    # No tabs, no tabstops — all plain left-aligned text
    glossary_lines = [
        "API Application Programming Interface specification",
        "CPU Central Processing Unit hardware component",
        "RAM Random Access Memory volatile storage",
        "GPU Graphics Processing Unit parallel processor",
        "DNS Domain Name System address resolution protocol",
        "HTTP HyperText Transfer Protocol web communication standard",
        "JSON JavaScript Object Notation lightweight data format",
        "SSH Secure Shell encrypted remote access protocol",
        "SQL Structured Query Language relational database interface",
        "TCP Transmission Control Protocol reliable data transport",
        "UDP User Datagram Protocol connectionless data transport",
        "URL Uniform Resource Locator web address identifier",
        "XML Extensible Markup Language structured data encoding",
        "VPN Virtual Private Network encrypted tunnel technology",
        "IDE Integrated Development Environment programming toolkit",
        "SDK Software Development Kit platform programming library",
        "ORM Object Relational Mapper database abstraction layer",
        "CLI Command Line Interface text-based user interaction",
    ]

    # Remove default empty paragraph if present
    default_para = doc.paragraphs[0] if doc.paragraphs else None

    for line in glossary_lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(11)
        # No tabstops added — paragraph_format.tab_stops is intentionally empty
        # No tab characters in the text

    # Remove the initial empty paragraph added by Document()
    if default_para is not None and default_para.text == '':
        p = default_para._element
        p.getparent().remove(p)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
