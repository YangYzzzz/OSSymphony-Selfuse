"""
Reward Script: Workplace Investigation Report Template
Task ID: writer_hr_070
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): All 8 required sections present
  Component 2 (0.20): Evidence List table with 4 columns
  Component 3 (0.20): At least 3 witness interview subsections
  Component 4 (0.15): Numbered findings of fact
  Component 5 (0.15): Line numbering enabled in section properties
"""

import os
import re
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_070'


def persist_app_state(domain):
    """Save any unsaved LibreOffice changes via Ctrl+S."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraph text (lowered) for section detection
    all_text = [p.text.strip().lower() for p in doc.paragraphs]
    all_text_raw = [p.text.strip() for p in doc.paragraphs]

    # =========================================================================
    # Component 1: All 8 required sections present (0.30 points)
    # Task requires: Case Information, Complaint Summary, Evidence List,
    # Witness Interviews, Findings of Fact, Legal Analysis, Conclusions,
    # Recommended Actions
    # =========================================================================
    try:
        required_sections = [
            ("case information", ["case information"]),
            ("complaint summary", ["complaint summary"]),
            ("evidence list", ["evidence list", "evidence"]),
            ("witness interview", ["witness interview"]),
            ("findings of fact", ["findings of fact", "findings"]),
            ("legal analysis", ["legal analysis"]),
            ("conclusions", ["conclusion"]),
            ("recommended actions", ["recommended action"]),
        ]

        sections_found = 0
        for section_name, keywords in required_sections:
            found = False
            for t in all_text:
                if any(kw in t for kw in keywords):
                    found = True
                    break
            if found:
                sections_found += 1
            else:
                print(f"  MISS: Section '{section_name}' not found")

        if sections_found == 8:
            print(f"PASS: Component 1 -- All 8 required sections found (0.30 pts)")
            total_score += 0.30
        elif sections_found >= 6:
            partial = round(0.30 * (sections_found / 8), 2)
            print(f"PARTIAL: Component 1 -- {sections_found}/8 sections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 -- Only {sections_found}/8 sections found")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # =========================================================================
    # Component 2: Evidence List table with 4 columns (0.20 points)
    # Task requires: table with columns item number, type, date collected,
    # description
    # =========================================================================
    try:
        evidence_table_found = False
        if len(doc.tables) > 0:
            for table in doc.tables:
                num_cols = len(table.columns)
                if num_cols >= 4:
                    # Check if header row has evidence-related column names
                    header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
                    header_text = ' '.join(header_cells)
                    # Look for evidence-table-like headers
                    has_item = any('item' in h or 'no' in h or '#' in h or 'number' in h for h in header_cells)
                    has_type = any('type' in h or 'category' in h for h in header_cells)
                    has_date = any('date' in h for h in header_cells)
                    has_desc = any('desc' in h for h in header_cells)

                    if (has_item or has_type) and num_cols >= 4:
                        evidence_table_found = True
                        # Check it has data rows (more than just header)
                        if len(table.rows) >= 2:
                            print(f"PASS: Component 2 -- Evidence table found: {num_cols} cols, {len(table.rows)} rows (0.20 pts)")
                            total_score += 0.20
                        else:
                            print(f"PARTIAL: Component 2 -- Evidence table found but only header row (0.10 pts)")
                            total_score += 0.10
                        break

            if not evidence_table_found:
                # Fallback: any table with >= 4 columns counts as partial
                for table in doc.tables:
                    if len(table.columns) >= 4:
                        print(f"PARTIAL: Component 2 -- Table with {len(table.columns)} cols found, but headers don't match evidence list pattern (0.10 pts)")
                        total_score += 0.10
                        evidence_table_found = True
                        break

            if not evidence_table_found:
                print(f"FAIL: Component 2 -- No table with 4+ columns found")
        else:
            print(f"FAIL: Component 2 -- No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # =========================================================================
    # Component 3: At least 3 witness interview subsections (0.20 points)
    # Each subsection should have a repeating format with witness name/number
    # =========================================================================
    try:
        witness_count = 0
        for t in all_text:
            # Match patterns like "witness 1", "witness 2", "witness #1", etc.
            if re.search(r'witness\s*[#]?\s*\d+', t):
                witness_count += 1
            # Also match "witness: [name]" or "witness name:" pattern (as heading)

        # Deduplicate: count unique witness numbers
        witness_numbers = set()
        for t in all_text:
            matches = re.findall(r'witness\s*[#]?\s*(\d+)', t)
            for m in matches:
                witness_numbers.add(int(m))

        # Also check for repeating subsection patterns (Name:, Date:, Summary, Credibility)
        credibility_count = sum(1 for t in all_text if 'credibility' in t)
        summary_statement_count = sum(1 for t in all_text if 'summary of statement' in t or 'summary' in t.split(':')[0] if 'statement' in t)

        num_witnesses = max(len(witness_numbers), min(credibility_count, summary_statement_count))

        if num_witnesses >= 3:
            print(f"PASS: Component 3 -- {num_witnesses} witness interview subsections found (0.20 pts)")
            total_score += 0.20
        elif num_witnesses >= 1:
            partial = round(0.20 * (num_witnesses / 3), 2)
            print(f"PARTIAL: Component 3 -- {num_witnesses}/3 witness subsections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 -- No witness interview subsections found")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # =========================================================================
    # Component 4: Numbered findings of fact (0.15 points)
    # Task requires numbered findings in the Findings section
    # =========================================================================
    try:
        # Find the findings section and look for numbered items
        in_findings = False
        numbered_findings = 0
        for t in all_text_raw:
            tl = t.lower()
            if 'findings of fact' in tl or 'findings' in tl.split('.')[0] if '.' in tl else tl == 'findings':
                in_findings = True
                continue
            # Stop when we hit the next major section
            if in_findings and re.match(r'^(vi|vii|viii|ix|x)[\.\s]', tl):
                break
            if in_findings and re.match(r'^(legal analysis|conclusion|recommended)', tl):
                break
            if in_findings:
                # Check for numbered items: "1.", "2.", etc.
                if re.match(r'^\d+[\.\)]\s', t):
                    numbered_findings += 1

        if numbered_findings >= 3:
            print(f"PASS: Component 4 -- {numbered_findings} numbered findings found (0.15 pts)")
            total_score += 0.15
        elif numbered_findings >= 1:
            partial = round(0.15 * (numbered_findings / 3), 2)
            print(f"PARTIAL: Component 4 -- {numbered_findings} numbered findings (need >= 3) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 -- No numbered findings found")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    # =========================================================================
    # Component 5: Line numbering enabled (0.15 points)
    # Task requires "line numbers displayed in the left margin"
    # Check w:lnNumType element in section properties
    # =========================================================================
    try:
        line_numbering_found = False
        for section in doc.sections:
            sectPr = section._sectPr
            lnNumType = sectPr.find(qn('w:lnNumType'))
            if lnNumType is not None:
                line_numbering_found = True
                break

        if line_numbering_found:
            print(f"PASS: Component 5 -- Line numbering enabled in document (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 -- No line numbering (w:lnNumType) found in section properties")
    except Exception as e:
        print(f"ERROR: Component 5 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
