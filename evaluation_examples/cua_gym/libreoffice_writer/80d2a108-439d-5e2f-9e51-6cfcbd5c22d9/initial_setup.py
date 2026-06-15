"""
Initial Setup: Invoice document without table (agent must create the table)
Task ID: writer_tbl_066
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_066'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DESKTOP = f'{WORKDIR}/Desktop'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title: INVOICE — bold, 18pt
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_run = title_para.add_run("INVOICE")
    title_run.bold = True
    title_run.font.size = Pt(18)

    # Invoice header metadata
    doc.add_paragraph("Invoice #: INV-2024-001")
    doc.add_paragraph("Date: March 15, 2024")
    doc.add_paragraph("Client: Acme Corporation")

    # NOTE: No table is created here — the agent must create the table

    # Save to home directory
    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # Also copy to Desktop so task context matches
    os.makedirs(DESKTOP, exist_ok=True)
    import shutil
    desktop_path = f'{DESKTOP}/{TASK_ID}.docx'
    shutil.copy(OUTPUT, desktop_path)
    print(f"Copied to Desktop: {desktop_path}")

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
