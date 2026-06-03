"""
Reward Script: Create an invoice table with proper formatting
Task ID: writer_tbl_066
Domain: libreoffice_writer
Scoring:
  Component 1: Table structure exists (6 rows x 5 cols)   — 0.3 pts
  Component 2: Header row is correct (all 5 column names)  — 0.2 pts
  Component 3: 4 product data rows are correct             — 0.3 pts
  Component 4: Grand Total row with merged cells + value    — 0.2 pts
  Total: 1.0
"""

import os
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_066'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: Table exists with correct dimensions (6 rows x 5 columns)
    #              This verifies the core task action of creating the table.
    # -----------------------------------------------------------------------
    try:
        tables = doc.tables
        if len(tables) < 1:
            print("FAIL: Component 1 — No table found in document")
        else:
            table = tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 6 and num_cols == 5:
                print(f"PASS: Component 1 — Table has correct structure: {num_rows} rows x {num_cols} cols (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 1 — Expected 6 rows x 5 cols, found {num_rows} rows x {num_cols} cols")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # If no table, skip remaining checks
    if len(doc.tables) < 1:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    table = doc.tables[0]

    # -----------------------------------------------------------------------
    # Component 2: Header row correct — all 5 expected column names present
    # -----------------------------------------------------------------------
    try:
        expected_headers = ['Item', 'Description', 'Qty', 'Unit Price', 'Total']
        header_row = table.rows[0]
        actual_headers = [cell.text.strip() for cell in header_row.cells]
        if actual_headers == expected_headers:
            print(f"PASS: Component 2 — Header row matches expected: {actual_headers} (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 2 — Expected headers {expected_headers}, found {actual_headers}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: 4 product data rows correct
    #   Row 2: 001 | Consulting Hours  | 10 | 150 | 1500
    #   Row 3: 002 | Software License  |  2 | 500 | 1000
    #   Row 4: 003 | Training Session  |  3 | 200 |  600
    #   Row 5: 004 | Support Package   |  1 | 800 |  800
    # -----------------------------------------------------------------------
    try:
        expected_data_rows = [
            ['001', 'Consulting Hours', '10', '150', '1500'],
            ['002', 'Software License', '2', '500', '1000'],
            ['003', 'Training Session', '3', '200', '600'],
            ['004', 'Support Package', '1', '800', '800'],
        ]
        rows_correct = 0
        rows_total = len(expected_data_rows)
        for i, expected_row in enumerate(expected_data_rows):
            actual_row = [cell.text.strip() for cell in table.rows[i + 1].cells]
            if actual_row == expected_row:
                rows_correct += 1
            else:
                print(f"FAIL: Component 3 — Row {i+2} mismatch. Expected {expected_row}, found {actual_row}")
        if rows_correct == rows_total:
            print(f"PASS: Component 3 — All 4 product data rows correct (0.3 pts)")
            total_score += 0.3
        elif rows_correct > 0:
            partial = round(0.3 * rows_correct / rows_total, 2)
            print(f"PARTIAL: Component 3 — {rows_correct}/{rows_total} data rows correct ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No data rows match expected content")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: Grand Total row
    #   - Cells A6:D6 (indices 0-3) merged into one cell with gridSpan=4
    #   - Text "Grand Total:" right-aligned
    #   - Cell E6 (index 4) contains "3900"
    # -----------------------------------------------------------------------
    try:
        last_row = table.rows[5]

        # Check merged cell: first cell should have gridSpan=4
        merged_cell = last_row.cells[0]
        tc_pr = merged_cell._tc.find(qn('w:tcPr'))
        grid_span_elem = tc_pr.find(qn('w:gridSpan')) if tc_pr is not None else None
        grid_span_val = int(grid_span_elem.get(qn('w:val'))) if grid_span_elem is not None else 0

        # Check text content of merged cell
        grand_total_text = merged_cell.text.strip()

        # Check right-alignment of the Grand Total cell
        is_right_aligned = any(
            para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT
            for para in merged_cell.paragraphs
        )

        # Check total value in column E (index 4)
        total_cell = last_row.cells[4]
        total_value = total_cell.text.strip()

        # Unique cells: merged cells appear as the same cell object
        seen_ids = set()
        unique_cells_count = 0
        for cell in last_row.cells:
            cell_id = id(cell._tc)
            if cell_id not in seen_ids:
                seen_ids.add(cell_id)
                unique_cells_count += 1

        sub_checks = []
        # Check 1: merged with gridSpan=4 (or 2 unique cells meaning cols 0-3 merged)
        if grid_span_val == 4 or unique_cells_count == 2:
            sub_checks.append(True)
            print(f"  Component 4 sub-check: Merged cell spans 4 columns (gridSpan={grid_span_val}, unique_cells={unique_cells_count})")
        else:
            sub_checks.append(False)
            print(f"  Component 4 sub-check FAIL: Expected gridSpan=4, found gridSpan={grid_span_val} (unique_cells={unique_cells_count})")

        # Check 2: text is "Grand Total:"
        if 'Grand Total' in grand_total_text and grand_total_text.rstrip().endswith(':'):
            sub_checks.append(True)
            print(f"  Component 4 sub-check: Grand Total text correct: {repr(grand_total_text)}")
        else:
            sub_checks.append(False)
            print(f"  Component 4 sub-check FAIL: Expected 'Grand Total:', found {repr(grand_total_text)}")

        # Check 3: right-aligned
        if is_right_aligned:
            sub_checks.append(True)
            print(f"  Component 4 sub-check: Grand Total cell is right-aligned")
        else:
            sub_checks.append(False)
            print(f"  Component 4 sub-check FAIL: Grand Total cell is not right-aligned (alignment={merged_cell.paragraphs[0].paragraph_format.alignment if merged_cell.paragraphs else None})")

        # Check 4: total value is 3900
        if total_value == '3900':
            sub_checks.append(True)
            print(f"  Component 4 sub-check: Total value is {repr(total_value)}")
        else:
            sub_checks.append(False)
            print(f"  Component 4 sub-check FAIL: Expected '3900', found {repr(total_value)}")

        passed_sub = sum(sub_checks)
        if passed_sub == 4:
            print(f"PASS: Component 4 — Grand Total row fully correct (0.2 pts)")
            total_score += 0.2
        elif passed_sub >= 2:
            partial = round(0.2 * passed_sub / 4, 2)
            print(f"PARTIAL: Component 4 — {passed_sub}/4 sub-checks passed ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Grand Total row incorrect ({passed_sub}/4 sub-checks passed)")

    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
