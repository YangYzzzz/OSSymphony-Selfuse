"""
Initial Setup: Disable AutoCorrect capitalize-first-letter feature
Task ID: writer_edit_026
Domain: libreoffice_writer

Creates code_samples.docx on the Desktop with realistic programming code snippets
mixed with explanatory text. AutoCorrect is in its default state (enabled).
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_026'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/code_samples.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading('Python Code Samples — Developer Reference Guide', level=0)

    # Introduction section
    doc.add_heading('Introduction', level=1)
    intro = doc.add_paragraph(
        'This document contains annotated Python code snippets for common design patterns. '
        'each example demonstrates best practices along with explanatory notes. '
        'the code is structured for readability and maintainability.'
    )

    # Section 1: Object-Oriented Patterns
    doc.add_heading('1. Object-Oriented Patterns', level=1)

    doc.add_paragraph(
        'object-oriented programming (OOP) is a paradigm that structures software around '
        'objects rather than functions and logic. the following examples show how to design '
        'clean, reusable class hierarchies in Python.'
    )

    doc.add_heading('1.1 Singleton Pattern', level=2)
    doc.add_paragraph(
        'the singleton pattern ensures a class has only one instance. '
        'obj.method() is often called on the singleton to perform operations. '
        'note: in Python, singletons can be implemented using class variables or metaclasses.'
    )

    # Code block (monospace)
    code_para1 = doc.add_paragraph()
    run = code_para1.add_run(
        'class DatabaseConnection:\n'
        '    _instance = None\n\n'
        '    @classmethod\n'
        '    def get_instance(cls):\n'
        '        if cls._instance is None:\n'
        '            cls._instance = cls()\n'
        '        return cls._instance\n\n'
        '    def query(self, sql):\n'
        '        # obj.method() pattern used throughout\n'
        '        return self._execute(sql)\n\n'
        '    def _execute(self, sql):\n'
        '        print(f"Executing: {sql}")\n'
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_paragraph(
        'usage example: db = DatabaseConnection.get_instance(). '
        'obj.connect() establishes the database link. '
        'calling obj.query("SELECT * FROM users") retrieves all user records.'
    )

    doc.add_heading('1.2 Observer Pattern', level=2)
    doc.add_paragraph(
        'the observer pattern defines a one-to-many dependency between objects. '
        'when obj.state changes, all dependents are notified automatically. '
        'this is widely used in event-driven systems and GUI frameworks.'
    )

    code_para2 = doc.add_paragraph()
    run2 = code_para2.add_run(
        'class EventEmitter:\n'
        '    def __init__(self):\n'
        '        self._listeners = {}\n\n'
        '    def on(self, event, callback):\n'
        '        # obj.on("click", handler) registers event\n'
        '        self._listeners.setdefault(event, []).append(callback)\n\n'
        '    def emit(self, event, *args):\n'
        '        for cb in self._listeners.get(event, []):\n'
        '            cb(*args)\n'
    )
    run2.font.name = 'Courier New'
    run2.font.size = Pt(10)

    # Section 2: Functional Programming
    doc.add_heading('2. Functional Programming Techniques', level=1)

    doc.add_paragraph(
        'functional programming treats computation as the evaluation of mathematical functions. '
        'python supports functional style through first-class functions, lambda expressions, '
        'and higher-order functions like map(), filter(), and reduce().'
    )

    doc.add_heading('2.1 Decorators', level=2)
    doc.add_paragraph(
        'a decorator modifies the behavior of a function without changing its source code. '
        'obj.decorated_method() will execute the wrapper logic before and after the original. '
        'decorators are widely used for logging, authentication, and caching.'
    )

    code_para3 = doc.add_paragraph()
    run3 = code_para3.add_run(
        'import functools\nimport time\n\n'
        'def timer(func):\n'
        '    """measure execution time of any function."""\n'
        '    @functools.wraps(func)\n'
        '    def wrapper(*args, **kwargs):\n'
        '        start = time.perf_counter()\n'
        '        result = func(*args, **kwargs)\n'
        '        elapsed = time.perf_counter() - start\n'
        '        # obj.timer wraps target function\n'
        '        print(f"{func.__name__} took {elapsed:.4f}s")\n'
        '        return result\n'
        '    return wrapper\n\n'
        '@timer\n'
        'def compute_primes(limit):\n'
        '    """return all primes up to limit using sieve."""\n'
        '    sieve = [True] * (limit + 1)\n'
        '    sieve[0] = sieve[1] = False\n'
        '    for i in range(2, int(limit**0.5) + 1):\n'
        '        if sieve[i]:\n'
        '            for j in range(i*i, limit+1, i):\n'
        '                sieve[j] = False\n'
        '    return [i for i, p in enumerate(sieve) if p]\n'
    )
    run3.font.name = 'Courier New'
    run3.font.size = Pt(10)

    # Section 3: Error Handling
    doc.add_heading('3. Error Handling Best Practices', level=1)

    doc.add_paragraph(
        'robust error handling is essential for production-quality software. '
        'obj.risky_operation() should always be wrapped in a try/except block. '
        'the following examples demonstrate structured exception handling patterns.'
    )

    code_para4 = doc.add_paragraph()
    run4 = code_para4.add_run(
        'class APIClient:\n'
        '    def __init__(self, base_url):\n'
        '        self.base_url = base_url\n'
        '        self.session = None\n\n'
        '    def fetch(self, endpoint):\n'
        '        # obj.fetch("/api/data") retrieves endpoint\n'
        '        try:\n'
        '            response = self.session.get(f"{self.base_url}/{endpoint}")\n'
        '            response.raise_for_status()\n'
        '            return response.json()\n'
        '        except ConnectionError as e:\n'
        '            print(f"network error: {e}")\n'
        '            return None\n'
        '        except ValueError:\n'
        '            print("invalid json response")\n'
        '            return {}\n'
    )
    run4.font.name = 'Courier New'
    run4.font.size = Pt(10)

    doc.add_paragraph(
        'note: always catch specific exceptions rather than bare except clauses. '
        'obj.cleanup() should be called in a finally block to release resources. '
        'context managers (with statement) are preferred for resource management.'
    )

    # Section 4: Data Structures
    doc.add_heading('4. Custom Data Structures', level=1)

    doc.add_paragraph(
        'python\'s built-in data structures (list, dict, set, tuple) cover most use cases. '
        'however, specialized structures like linked lists, trees, and heaps offer '
        'performance advantages for specific algorithms. obj.insert() and obj.search() '
        'are the core operations for most tree implementations.'
    )

    code_para5 = doc.add_paragraph()
    run5 = code_para5.add_run(
        'class BinarySearchTree:\n'
        '    class Node:\n'
        '        def __init__(self, value):\n'
        '            self.value = value\n'
        '            self.left = self.right = None\n\n'
        '    def __init__(self):\n'
        '        self.root = None\n\n'
        '    def insert(self, value):\n'
        '        # obj.insert(42) adds node to tree\n'
        '        self.root = self._insert(self.root, value)\n\n'
        '    def _insert(self, node, value):\n'
        '        if node is None:\n'
        '            return self.Node(value)\n'
        '        if value < node.value:\n'
        '            node.left = self._insert(node.left, value)\n'
        '        elif value > node.value:\n'
        '            node.right = self._insert(node.right, value)\n'
        '        return node\n'
    )
    run5.font.name = 'Courier New'
    run5.font.size = Pt(10)

    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'these patterns form the foundation of professional Python development. '
        'obj.apply() of these techniques consistently leads to maintainable, testable code. '
        'refer to the official Python documentation and PEP guidelines for additional guidance. '
        'the examples above can be adapted for use in web applications, data pipelines, '
        'and system automation scripts.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
