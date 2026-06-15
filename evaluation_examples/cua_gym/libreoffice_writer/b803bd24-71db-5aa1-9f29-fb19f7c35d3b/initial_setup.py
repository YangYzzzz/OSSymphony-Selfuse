"""
Initial Setup: Reply to existing comment on 'methodology' in journal paper
Task ID: writer_struct_023
Domain: libreoffice_writer

Creates a 5-page journal paper (journal_submission.docx) at ~/Desktop/ with:
- ONE comment on the word 'methodology' in the first paragraph
- Comment text: "Please clarify which methodology was used."
- No replies to the comment (task is to add one)
"""

import os
import shlex
import subprocess
import time
import shutil
import zipfile
import copy
import io
from lxml import etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_struct_023'
FILENAME = 'journal_submission.docx'
OUTPUT = f'{WORKDIR}/{FILENAME}'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14_NS = 'http://schemas.microsoft.com/office/word/2010/wordml'
W15_NS = 'http://schemas.microsoft.com/office/word/2012/wordml'
R_NS_URL = 'http://schemas.openxmlformats.org/package/2006/relationships'
CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
COMMENTS_RT = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_run_element(text, rPr=None):
    """Create a <w:r> element with given text."""
    r = etree.Element('{%s}r' % W_NS)
    if rPr is not None:
        r.append(copy.deepcopy(rPr))
    t = etree.SubElement(r, '{%s}t' % W_NS)
    t.text = text
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r


def make_comment_range_start(cid):
    el = etree.Element('{%s}commentRangeStart' % W_NS)
    el.set('{%s}id' % W_NS, str(cid))
    return el


def make_comment_range_end(cid):
    el = etree.Element('{%s}commentRangeEnd' % W_NS)
    el.set('{%s}id' % W_NS, str(cid))
    return el


def make_comment_reference(cid):
    r = etree.Element('{%s}r' % W_NS)
    rPr_el = etree.SubElement(r, '{%s}rPr' % W_NS)
    rStyle = etree.SubElement(rPr_el, '{%s}rStyle' % W_NS)
    rStyle.set('{%s}val' % W_NS, 'CommentReference')
    ref = etree.SubElement(r, '{%s}commentReference' % W_NS)
    ref.set('{%s}id' % W_NS, str(cid))
    return r


def build_base_document_bytes():
    """Build the base 5-page journal paper using python-docx and return as bytes."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Title
    title_para = doc.add_heading(
        "Mixed-Methods Research on Digital Literacy in Higher Education: A Comprehensive Analysis",
        level=1
    )

    # First paragraph (with 'methodology' — will receive comment)
    first_para = doc.add_paragraph()
    first_para.add_run("Abstract: ")
    first_para.add_run(
        "This study employed a rigorous methodology to investigate the research questions "
        "outlined in Section 2. The research was conducted over a period of eighteen months, "
        "involving participants from three universities in Southeast Asia."
    )

    # Introduction
    doc.add_heading("1. Introduction", level=2)
    doc.add_paragraph(
        "Digital literacy has emerged as a critical competency in contemporary higher education "
        "settings. As institutions increasingly integrate technology into their pedagogical "
        "frameworks, understanding how students develop and utilize digital skills becomes "
        "paramount for educational policy and practice (Smith et al., 2022)."
    )
    doc.add_paragraph(
        "Previous studies have documented the uneven distribution of digital competencies across "
        "demographic groups (Johnson & Williams, 2021). However, few have examined the longitudinal "
        "development of these skills within institutional contexts. This gap in the literature "
        "motivates the present investigation."
    )

    # Research Questions
    doc.add_heading("2. Research Questions", level=2)
    doc.add_paragraph("The present study addresses three primary research questions:")
    doc.add_paragraph(
        "(1) How do students' digital literacy skills develop throughout their undergraduate education?",
        style="List Number"
    )
    doc.add_paragraph(
        "(2) What institutional factors mediate this development?",
        style="List Number"
    )
    doc.add_paragraph(
        "(3) How do socioeconomic background and prior technology exposure influence learning trajectories?",
        style="List Number"
    )

    doc.add_page_break()

    # Methodology Section
    doc.add_heading("3. Methodology", level=2)
    doc.add_paragraph(
        "This research employs a mixed-methods approach, integrating both quantitative survey data "
        "and qualitative interview narratives. The quantitative component involved a validated digital "
        "literacy assessment instrument administered to 847 undergraduate students across three cohorts. "
        "The qualitative component comprised semi-structured interviews with 45 purposively selected participants."
    )
    doc.add_paragraph(
        "Data collection occurred at four timepoints: semester 1 (baseline), semester 3, semester 5, "
        "and semester 7 (final assessment). This longitudinal design enables tracking of skill development "
        "trajectories while controlling for curriculum effects."
    )
    doc.add_paragraph(
        "Statistical analyses employed hierarchical linear modeling to account for nested data structures "
        "(students within institutions). Interview data were analyzed using thematic analysis following "
        "the framework of Braun and Clarke (2019), with inter-rater reliability established at κ = 0.82."
    )

    doc.add_page_break()

    # Results
    doc.add_heading("4. Results", level=2)
    doc.add_paragraph("4.1 Quantitative Findings")
    doc.add_paragraph(
        "Significant improvements in digital literacy scores were observed from baseline to final "
        "assessment (M baseline = 58.3, SD = 12.4; M final = 74.7, SD = 10.9; t(846) = 28.4, p < .001, "
        "d = 1.42). This represents a substantial effect size indicating meaningful skill development "
        "over the undergraduate period."
    )
    doc.add_paragraph("4.2 Qualitative Findings")
    doc.add_paragraph(
        "Thematic analysis revealed four primary themes: (1) Technology integration anxiety, "
        "(2) Peer learning networks, (3) Institutional support structures, and (4) Intrinsic motivation "
        "and self-efficacy. Students consistently reported that collaborative project-based learning "
        "experiences were the most transformative for their digital skill development."
    )

    doc.add_page_break()

    # Discussion
    doc.add_heading("5. Discussion", level=2)
    doc.add_paragraph(
        "The convergence of quantitative and qualitative findings supports a holistic interpretation of "
        "digital literacy development as a socially-mediated process. Students do not develop these "
        "competencies in isolation but rather through structured interactions with peers, instructors, "
        "and institutional resources."
    )
    doc.add_paragraph(
        "These findings have important implications for curriculum design and faculty development "
        "programs. Institutions seeking to enhance digital literacy outcomes should prioritize "
        "collaborative learning environments and provide structured scaffolding for technology "
        "integration activities."
    )

    # Conclusion
    doc.add_heading("6. Conclusion", level=2)
    doc.add_paragraph(
        "This study makes three primary contributions to the literature on digital literacy in higher "
        "education. First, it provides longitudinal evidence of skill development trajectories, addressing "
        "a significant gap in cross-sectional research. Second, it identifies key institutional mediating "
        "factors that can inform policy decisions. Third, it demonstrates the value of mixed-methods "
        "designs for capturing the complexity of digital literacy development."
    )
    doc.add_paragraph(
        "Future research should examine how specific pedagogical interventions can accelerate digital "
        "literacy development among students with limited prior technology exposure. Comparative "
        "international studies would also contribute valuable cross-cultural perspectives to this "
        "emerging body of knowledge."
    )

    # References
    doc.add_heading("References", level=2)
    doc.add_paragraph(
        "Braun, V., & Clarke, V. (2019). Thematic analysis: A practical guide. SAGE Publications."
    )
    doc.add_paragraph(
        "Johnson, M., & Williams, K. (2021). Digital divides in higher education: A systematic review. "
        "Journal of Educational Technology Research, 15(3), 112-134."
    )
    doc.add_paragraph(
        "Smith, A., Torres, R., & Chen, L. (2022). Digital literacy frameworks for the 21st century. "
        "International Review of Education, 68(4), 521-547."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def inject_comment_into_docx(docx_bytes):
    """
    Inject a single comment on the word 'methodology' in paragraph 1 (abstract).
    Returns modified docx bytes with:
    - word/comments.xml added
    - word/_rels/document.xml.rels updated
    - [Content_Types].xml updated
    - word/document.xml updated with commentRangeStart/End/Reference markers
    """
    # Read all files from the docx bytes
    all_files = {}
    with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as z:
        for name in z.namelist():
            all_files[name] = z.read(name)

    # --- Modify document.xml ---
    doc_tree = etree.fromstring(all_files['word/document.xml'])
    body = doc_tree.find('.//{%s}body' % W_NS)
    paragraphs = body.findall('{%s}p' % W_NS)

    # Paragraph index 1 is the abstract paragraph containing 'methodology'
    target_para = paragraphs[1]
    runs = target_para.findall('{%s}r' % W_NS)

    # Run index 1 contains "This study employed a rigorous methodology..."
    target_run = runs[1]
    run_text_full = ''.join((t.text or '') for t in target_run.findall('{%s}t' % W_NS))
    idx = run_text_full.lower().find('methodology')
    before_text = run_text_full[:idx]        # "This study employed a rigorous "
    method_text = run_text_full[idx:idx+11]  # "methodology"
    after_text = run_text_full[idx+11:]      # " to investigate..."

    rPr = target_run.find('{%s}rPr' % W_NS)

    # Replace original run with: r_before + commentRangeStart + r_method +
    #   commentRangeEnd + commentReference + r_after
    parent = target_run.getparent()
    insert_idx = list(parent).index(target_run)
    parent.remove(target_run)

    new_elements = [
        make_run_element(before_text, rPr),
        make_comment_range_start(1),
        make_run_element(method_text, rPr),
        make_comment_range_end(1),
        make_comment_reference(1),
        make_run_element(after_text, rPr),
    ]
    for i, el in enumerate(new_elements):
        parent.insert(insert_idx + i, el)

    all_files['word/document.xml'] = etree.tostring(
        doc_tree, xml_declaration=True, encoding='UTF-8', standalone=True
    )

    # --- Create comments.xml (single comment, no replies) ---
    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:comments xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"\n'
        '            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"\n'
        '            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"\n'
        '            xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"\n'
        '            xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"\n'
        '            mc:Ignorable="w14 w15">\n'
        '  <w:comment w:id="1" w:author="Dr. Elena Vasquez" w:date="2025-10-15T09:30:00Z" w:initials="EV">\n'
        '    <w:p w14:paraId="3F7A1B2C" w14:textId="77777777">\n'
        '      <w:pPr><w:pStyle w:val="CommentText"/></w:pPr>\n'
        '      <w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:annotationRef/></w:r>\n'
        '      <w:r><w:t>Please clarify which methodology was used.</w:t></w:r>\n'
        '    </w:p>\n'
        '  </w:comment>\n'
        '</w:comments>'
    )
    all_files['word/comments.xml'] = comments_xml.encode('utf-8')

    # --- Update document.xml.rels ---
    rels_tree = etree.fromstring(all_files['word/_rels/document.xml.rels'])
    # Check if comments rel already exists
    existing = [
        el for el in rels_tree
        if el.get('Type') == COMMENTS_RT
    ]
    if not existing:
        rel = etree.SubElement(rels_tree, '{%s}Relationship' % R_NS_URL)
        rel.set('Id', 'rId_comments')
        rel.set('Type', COMMENTS_RT)
        rel.set('Target', 'comments.xml')
    all_files['word/_rels/document.xml.rels'] = etree.tostring(
        rels_tree, xml_declaration=True, encoding='UTF-8', standalone=True
    )

    # --- Update [Content_Types].xml ---
    ct_tree = etree.fromstring(all_files['[Content_Types].xml'])
    existing_ct = [el for el in ct_tree if el.get('PartName') == '/word/comments.xml']
    if not existing_ct:
        ov = etree.SubElement(ct_tree, '{%s}Override' % CT_NS)
        ov.set('PartName', '/word/comments.xml')
        ov.set(
            'ContentType',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
        )
    all_files['[Content_Types].xml'] = etree.tostring(
        ct_tree, xml_declaration=True, encoding='UTF-8', standalone=True
    )

    # --- Write updated docx ---
    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in all_files.items():
            zout.writestr(name, data)
    out_buf.seek(0)
    return out_buf.read()


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    # Build base document bytes
    base_bytes = build_base_document_bytes()

    # Inject comment on "methodology"
    final_bytes = inject_comment_into_docx(base_bytes)

    # Write to file
    with open(OUTPUT, 'wb') as f:
        f.write(final_bytes)

    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
