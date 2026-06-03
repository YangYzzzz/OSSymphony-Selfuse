"""
Initial Setup: Create a 6-page technical specification document with a uniform header
Task ID: writer_page_018
Domain: libreoffice_writer

Initial state:
  - File at /home/user/Desktop/tech_spec.docx
  - 6-page technical specification document
  - Page setup: A4, portrait, margins 2.54 cm all sides
  - Header enabled on ALL pages with text 'Technical Specification v2.1'
  - 'Different first page' header option is NOT enabled
  - LibreOffice Writer opens the file
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
FILENAME = 'tech_spec.docx'
OUTPUT = f'{WORKDIR}/{FILENAME}'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add a page break paragraph."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)
    return para


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set up page: A4 (21 cm x 29.7 cm), portrait, 2.54 cm margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Set uniform header (same on all pages including first page)
    # different_first_page_header_footer must be False (default)
    section.different_first_page_header_footer = False

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    run = hp.add_run('Technical Specification v2.1')
    run.font.size = Pt(11)
    hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # -------------------------
    # PAGE 1: Cover / Title
    # -------------------------
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    run = title_para.add_run('Technical Specification')
    run.bold = True
    run.font.size = Pt(24)

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle_para.add_run('Version 2.1')
    run.font.size = Pt(18)
    run.italic = True

    doc.add_paragraph()

    meta_para = doc.add_paragraph()
    meta_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta_para.add_run('Prepared by: Engineering Team')
    run.font.size = Pt(12)

    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Date: March 2025')
    run.font.size = Pt(12)

    doc.add_paragraph()

    classification_para = doc.add_paragraph()
    classification_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = classification_para.add_run('CONFIDENTIAL — INTERNAL USE ONLY')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    add_page_break(doc)

    # -------------------------
    # PAGE 2: Table of Contents
    # -------------------------
    toc_heading = doc.add_heading('Table of Contents', level=1)

    toc_entries = [
        ('1. Introduction', '3'),
        ('2. System Architecture', '4'),
        ('3. Data Flow Design', '5'),
        ('4. API Specifications', '5'),
        ('5. Security Considerations', '6'),
        ('6. Deployment Requirements', '6'),
        ('7. Testing Strategy', '7'),
        ('8. Appendix', '7'),
    ]
    for entry, page in toc_entries:
        toc_para = doc.add_paragraph()
        run_entry = toc_para.add_run(f'{entry}')
        run_entry.font.size = Pt(11)
        toc_para.add_run('\t')
        run_page = toc_para.add_run(page)
        run_page.font.size = Pt(11)

    add_page_break(doc)

    # -------------------------
    # PAGE 3: Introduction
    # -------------------------
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This document provides a comprehensive technical specification for the '
        'DataSync Platform v2.1, an enterprise-grade data synchronization solution '
        'designed to facilitate real-time data exchange between heterogeneous systems. '
        'The platform supports bidirectional synchronization with conflict resolution '
        'algorithms and guaranteed delivery semantics.'
    )
    doc.add_paragraph(
        'The system has been designed to handle up to 100,000 transactions per second '
        'with sub-millisecond latency under standard operating conditions. It leverages '
        'a distributed microservices architecture with automatic failover and horizontal '
        'scaling capabilities.'
    )

    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        'The primary purpose of this specification is to define the functional and '
        'non-functional requirements, architectural decisions, and implementation '
        'guidelines for the DataSync Platform. This document serves as the authoritative '
        'reference for development, testing, and deployment teams.'
    )

    doc.add_heading('1.2 Scope', level=2)
    doc.add_paragraph(
        'This specification covers the core synchronization engine, API layer, security '
        'framework, monitoring subsystem, and deployment infrastructure. It does not '
        'cover third-party integrations beyond the provided SDK interfaces.'
    )

    add_page_break(doc)

    # -------------------------
    # PAGE 4: System Architecture
    # -------------------------
    doc.add_heading('2. System Architecture', level=1)
    doc.add_paragraph(
        'The DataSync Platform employs a layered architecture consisting of five primary '
        'tiers: the Client Interface Layer, the API Gateway, the Core Processing Engine, '
        'the Data Store Manager, and the Monitoring & Alerting subsystem.'
    )

    doc.add_heading('2.1 Component Overview', level=2)
    components = [
        ('API Gateway', 'Routes and authenticates all inbound requests. Implements rate limiting, request validation, and load balancing across processing nodes.'),
        ('Core Processing Engine', 'Manages the synchronization lifecycle including conflict detection, resolution, and rollback. Operates on a queue-based messaging model using Apache Kafka.'),
        ('Data Store Manager', 'Abstracts the underlying storage layer. Supports PostgreSQL, MongoDB, and Redis adapters. Handles connection pooling and query optimization.'),
        ('Monitoring Subsystem', 'Provides real-time metrics via Prometheus and Grafana dashboards. Supports configurable alert thresholds and PagerDuty integration.'),
    ]
    for name, desc in components:
        comp_para = doc.add_paragraph(style='List Bullet')
        run_name = comp_para.add_run(f'{name}: ')
        run_name.bold = True
        run_name.font.size = Pt(11)
        run_desc = comp_para.add_run(desc)
        run_desc.font.size = Pt(11)

    doc.add_heading('2.2 Technology Stack', level=2)
    tech_stack = [
        'Backend: Python 3.11, FastAPI, Celery',
        'Message Broker: Apache Kafka 3.5',
        'Primary Database: PostgreSQL 15.2',
        'Cache Layer: Redis 7.2',
        'Container Orchestration: Kubernetes 1.28',
        'Service Mesh: Istio 1.19',
    ]
    for item in tech_stack:
        tech_para = doc.add_paragraph(item, style='List Bullet')

    add_page_break(doc)

    # -------------------------
    # PAGE 5: Data Flow & API
    # -------------------------
    doc.add_heading('3. Data Flow Design', level=1)
    doc.add_paragraph(
        'Data flows through the platform in discrete processing stages. Each stage is '
        'independently scalable and fault-tolerant. The following describes the primary '
        'synchronization workflow:'
    )

    steps = [
        'Client submits a synchronization request via the REST API or WebSocket connection.',
        'The API Gateway validates the request payload against the JSON Schema and authenticates the client credentials.',
        'A job record is created in the PostgreSQL database with status PENDING.',
        'The job is published to the Kafka topic `sync.requests` with a unique correlation ID.',
        'A worker node consumes the message, fetches the source data, and applies the configured transformation pipeline.',
        'Conflict detection algorithms compare source and target state vectors.',
        'Resolved changes are committed atomically to the target data store.',
        'Job status is updated to COMPLETED and a webhook notification is dispatched.',
    ]
    for i, step in enumerate(steps, 1):
        step_para = doc.add_paragraph(f'{i}. {step}')
        step_para.paragraph_format.left_indent = Cm(0.5)

    doc.add_heading('4. API Specifications', level=1)
    doc.add_paragraph(
        'The platform exposes a RESTful API conforming to OpenAPI 3.1 specification. '
        'Authentication uses OAuth 2.0 with JWT tokens. Base URL: https://api.datasync.internal/v2'
    )

    endpoints_table = doc.add_table(rows=5, cols=3)
    endpoints_table.style = 'Table Grid'
    headers_row = endpoints_table.rows[0]
    for cell, text in zip(headers_row.cells, ['Method', 'Endpoint', 'Description']):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    api_data = [
        ('POST', '/sync/jobs', 'Create a new synchronization job'),
        ('GET', '/sync/jobs/{id}', 'Retrieve job status and results'),
        ('DELETE', '/sync/jobs/{id}', 'Cancel a pending or running job'),
        ('GET', '/sync/metrics', 'Retrieve real-time platform metrics'),
    ]
    for i, (method, endpoint, desc) in enumerate(api_data, 1):
        row = endpoints_table.rows[i]
        row.cells[0].text = method
        row.cells[1].text = endpoint
        row.cells[2].text = desc

    add_page_break(doc)

    # -------------------------
    # PAGE 6: Security & Deployment
    # -------------------------
    doc.add_heading('5. Security Considerations', level=1)
    doc.add_paragraph(
        'All data transmitted through the platform must be encrypted using TLS 1.3 or '
        'higher. Data at rest is encrypted using AES-256-GCM. Encryption keys are managed '
        'via HashiCorp Vault with automatic rotation every 90 days.'
    )

    security_items = [
        'Authentication: OAuth 2.0 with PKCE flow for all API clients',
        'Authorization: Role-based access control (RBAC) with fine-grained permissions',
        'Audit Logging: All data access events logged to immutable audit trail',
        'Vulnerability Scanning: Automated weekly scans via Trivy and Snyk',
        'Penetration Testing: Quarterly assessments by certified security firm',
    ]
    for item in security_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6. Deployment Requirements', level=1)
    doc.add_paragraph(
        'Minimum production cluster configuration requires three Kubernetes worker nodes '
        'with 8 vCPUs and 32 GB RAM each. High availability configuration requires '
        'deployment across at least two availability zones.'
    )

    req_table = doc.add_table(rows=5, cols=3)
    req_table.style = 'Table Grid'
    req_headers = req_table.rows[0]
    for cell, text in zip(req_headers.cells, ['Component', 'Min Instances', 'Resources']):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    req_data = [
        ('API Gateway', '2', '2 vCPU, 4 GB RAM'),
        ('Processing Workers', '4', '4 vCPU, 8 GB RAM'),
        ('Database (Primary)', '1', '8 vCPU, 32 GB RAM'),
        ('Cache (Redis Cluster)', '3', '2 vCPU, 16 GB RAM'),
    ]
    for i, (comp, instances, resources) in enumerate(req_data, 1):
        row = req_table.rows[i]
        row.cells[0].text = comp
        row.cells[1].text = instances
        row.cells[2].text = resources

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
