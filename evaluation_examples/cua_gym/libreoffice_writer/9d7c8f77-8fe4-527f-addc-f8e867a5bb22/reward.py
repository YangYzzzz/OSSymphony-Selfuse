"""
Reward Script: Make the first page have a different header from the rest.
Task ID: writer_page_018
Domain: libreoffice_writer
Scoring:
  Component 1: 'Different First Page Header/Footer' is enabled (0.4 pts)
  Component 2: First page header text == 'Cover Page' (0.3 pts)
  Component 3: Main (subsequent pages) header text == 'Technical Specification v2.1'
               AND different_first_page is enabled (compound check) (0.3 pts)
  Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_page_018'
FILE_PATH = f'{WORKDIR}/Desktop/tech_spec.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.

    Task: Make the first page have a different header from the rest.
    - First page header should say 'Cover Page'
    - Subsequent pages header should say 'Technical Specification v2.1'
    - 'Different First Page Header' option must be enabled (titlePg/different_first_page_header_footer)

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if not doc.sections:
        print("CRITICAL: Document has no sections.")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: 'Different First Page Header/Footer' is enabled (0.4 points)
    # This is the 'Same content on first page' option being unchecked — i.e. titlePg is present.
    # On initial_env: different_first_page_header_footer == False → FAIL
    # On golden_env:  different_first_page_header_footer == True  → PASS
    try:
        diff_first_page = section.different_first_page_header_footer
        if diff_first_page:
            print("PASS: Component 1 — 'Different First Page Header/Footer' is enabled (0.4 pts)")
            total_score += 0.4
        else:
            print("FAIL: Component 1 — 'Different First Page Header/Footer' is NOT enabled; "
                  "expected True, found False")
    except Exception as e:
        print(f"ERROR: Component 1 — Could not check different_first_page_header_footer: {e}")

    # Component 2: First page header text is 'Cover Page' (0.3 points)
    # On initial_env: first_page_header text is '' → FAIL
    # On golden_env:  first_page_header text is 'Cover Page' → PASS
    try:
        first_page_hdr = section.first_page_header
        if first_page_hdr and first_page_hdr.paragraphs:
            fp_text = first_page_hdr.paragraphs[0].text.strip()
        else:
            fp_text = ""
        if fp_text == "Cover Page":
            print(f"PASS: Component 2 — First page header text is 'Cover Page' (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Expected first page header 'Cover Page', found: {repr(fp_text)}")
    except Exception as e:
        print(f"ERROR: Component 2 — Could not check first page header text: {e}")

    # Component 3: Main (subsequent pages) header text == 'Technical Specification v2.1'
    # AND different_first_page_header_footer is enabled (compound check) (0.3 points)
    #
    # This is a compound check: we require BOTH conditions to be true together.
    # The main header text exists on initial too, but without different_first_page enabled
    # the first page header would be the same — so this compound check ensures the task
    # is truly complete (different page headers are active AND subsequent pages are correct).
    #
    # On initial_env: different_first_page=False → compound check FAILS → 0 pts
    # On golden_env:  different_first_page=True AND main header='Technical Specification v2.1' → PASS
    try:
        main_hdr = section.header
        if main_hdr and main_hdr.paragraphs:
            main_text = main_hdr.paragraphs[0].text.strip()
        else:
            main_text = ""

        diff_first_page_check = section.different_first_page_header_footer
        main_text_correct = (main_text == "Technical Specification v2.1")

        if diff_first_page_check and main_text_correct:
            print(f"PASS: Component 3 — different_first_page enabled AND main header text is "
                  f"'Technical Specification v2.1' (0.3 pts)")
            total_score += 0.3
        elif not diff_first_page_check:
            print(f"FAIL: Component 3 — different_first_page not enabled; "
                  f"compound check requires both conditions (main header: {repr(main_text)})")
        else:
            print(f"FAIL: Component 3 — different_first_page is enabled but main header text is wrong: "
                  f"expected 'Technical Specification v2.1', found {repr(main_text)}")
    except Exception as e:
        print(f"ERROR: Component 3 — Could not check main header text: {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
