"""
Initial Setup: Create a 6-page business contract document without watermark
Task ID: writer_biz_046
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_046'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    return h


def add_body(doc, text, bold=False, alignment=None, space_after=Pt(6)):
    """Add a body paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    if alignment:
        para.paragraph_format.alignment = alignment
    para.paragraph_format.space_after = space_after
    return para


def create_initial():
    doc = Document()

    # Set default page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ========== PAGE 1: Title and Parties ==========
    # Title
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_body(doc, '')
    add_body(doc, 'Agreement Number: PSA-2025-04782', bold=True,
             alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_body(doc, 'Effective Date: March 15, 2025',
             alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_body(doc, '')

    add_heading_styled(doc, '1. PARTIES', level=1)

    add_body(doc, 'This Professional Services Agreement ("Agreement") is entered into as of '
             'March 15, 2025 ("Effective Date"), by and between:')
    add_body(doc, '')

    add_body(doc, 'Meridian Technologies Inc., a Delaware corporation with its principal '
             'place of business at 4200 Innovation Boulevard, Suite 800, San Jose, CA 95134 '
             '("Client");', bold=False)
    add_body(doc, '')
    add_body(doc, 'and')
    add_body(doc, '')
    add_body(doc, 'NexGen Solutions Group LLC, a California limited liability company with '
             'its principal place of business at 1750 Gateway Drive, Floor 12, Oakland, CA 94612 '
             '("Service Provider").')
    add_body(doc, '')

    add_body(doc, 'Client and Service Provider are each referred to herein as a "Party" and '
             'collectively as the "Parties."')

    add_heading_styled(doc, '2. BACKGROUND AND PURPOSE', level=1)

    add_body(doc, 'WHEREAS, Client desires to engage Service Provider to perform certain '
             'professional consulting, software development, and systems integration services '
             'in connection with Client\'s enterprise resource planning modernization initiative '
             '(the "Project");')
    add_body(doc, '')
    add_body(doc, 'WHEREAS, Service Provider possesses the requisite expertise, personnel, '
             'and resources to perform such services;')
    add_body(doc, '')
    add_body(doc, 'NOW, THEREFORE, in consideration of the mutual covenants and agreements '
             'set forth herein, and for other good and valuable consideration, the receipt and '
             'sufficiency of which are hereby acknowledged, the Parties agree as follows.')

    # ========== PAGE 2: Scope of Services ==========
    doc.add_page_break()

    add_heading_styled(doc, '3. SCOPE OF SERVICES', level=1)

    add_body(doc, '3.1 Service Provider shall provide the following professional services '
             '("Services") to Client:')
    add_body(doc, '')

    services = [
        '(a) Assessment and analysis of Client\'s existing ERP infrastructure, including '
        'Oracle E-Business Suite 12.2, SAP Business One integration points, and legacy '
        'mainframe data feeds;',
        '(b) Design and architecture of a cloud-native ERP solution leveraging Microsoft '
        'Dynamics 365 Finance and Operations, with custom modules for supply chain visibility '
        'and advanced financial reporting;',
        '(c) Data migration planning and execution, including extraction, transformation, '
        'and loading of approximately 14.2 million transaction records spanning fiscal years '
        '2018 through 2024;',
        '(d) Custom development of three integration APIs to connect the new ERP platform '
        'with Client\'s existing CRM (Salesforce Enterprise), warehouse management system '
        '(Manhattan Associates), and business intelligence platform (Tableau Server);',
        '(e) User acceptance testing coordination, including preparation of test scripts, '
        'defect tracking, and regression testing across all functional modules;',
        '(f) End-user training program development and delivery for approximately 340 users '
        'across six regional offices, including role-based training curricula and reference documentation.',
    ]
    for s in services:
        p = doc.add_paragraph(s, style='List Bullet')

    add_body(doc, '')
    add_body(doc, '3.2 The Services shall be performed in accordance with the project timeline '
             'and milestones set forth in Exhibit A (Statement of Work), which is incorporated '
             'herein by reference.')

    add_body(doc, '3.3 Any changes to the scope of Services shall require a written change '
             'order signed by authorized representatives of both Parties, in accordance with '
             'the change management procedures described in Section 11.')

    # ========== PAGE 3: Compensation and Payment ==========
    doc.add_page_break()

    add_heading_styled(doc, '4. COMPENSATION AND PAYMENT TERMS', level=1)

    add_body(doc, '4.1 Fixed Fee. In consideration for the Services, Client shall pay Service '
             'Provider a total fixed fee of Two Million Four Hundred Seventy-Five Thousand '
             'Dollars ($2,475,000.00) ("Fee"), payable in accordance with the milestone-based '
             'payment schedule set forth below:')
    add_body(doc, '')

    # Payment schedule table
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    headers = ['Milestone', 'Description', 'Due Date', 'Amount']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    payments = [
        ['M1', 'Project Kickoff & Assessment', 'April 1, 2025', '$371,250.00'],
        ['M2', 'Architecture Sign-off', 'June 15, 2025', '$495,000.00'],
        ['M3', 'Data Migration Complete', 'September 30, 2025', '$495,000.00'],
        ['M4', 'UAT Sign-off', 'December 15, 2025', '$495,000.00'],
        ['M5', 'Go-Live', 'February 28, 2026', '$371,250.00'],
        ['M6', 'Post-Go-Live Support Complete', 'April 30, 2026', '$247,500.00'],
    ]
    for r, row_data in enumerate(payments, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    add_body(doc, '')
    add_body(doc, '4.2 Payment Terms. All invoices shall be due and payable within thirty (30) '
             'calendar days of Client\'s receipt of a properly submitted invoice. Late payments '
             'shall accrue interest at the rate of 1.5% per month or the maximum rate permitted '
             'by applicable law, whichever is less.')

    add_body(doc, '4.3 Expenses. Service Provider shall be entitled to reimbursement for '
             'reasonable, pre-approved travel and out-of-pocket expenses incurred in connection '
             'with the Services, not to exceed $125,000.00 in aggregate without prior written '
             'approval from Client\'s Project Director.')

    # ========== PAGE 4: Confidentiality ==========
    doc.add_page_break()

    add_heading_styled(doc, '5. CONFIDENTIALITY', level=1)

    add_body(doc, '5.1 Definition. "Confidential Information" means all non-public information '
             'disclosed by either Party to the other, whether orally, in writing, or by any other '
             'means, that is designated as confidential or that reasonably should be understood '
             'to be confidential given the nature of the information and the circumstances of '
             'disclosure. Confidential Information includes, without limitation:')
    add_body(doc, '')

    conf_items = [
        '(a) Trade secrets, proprietary algorithms, source code, object code, and software architecture;',
        '(b) Business plans, financial projections, pricing strategies, and customer lists;',
        '(c) Technical specifications, system configurations, and network architecture diagrams;',
        '(d) Employee data, compensation structures, and organizational information;',
        '(e) Any materials marked or identified as "Confidential," "Proprietary," or similar designation.',
    ]
    for item in conf_items:
        doc.add_paragraph(item, style='List Bullet')

    add_body(doc, '')
    add_body(doc, '5.2 Obligations. Each Party shall: (i) hold all Confidential Information '
             'in strict confidence; (ii) not disclose Confidential Information to any third party '
             'without the prior written consent of the disclosing Party; (iii) use Confidential '
             'Information solely for the purposes of performing its obligations under this '
             'Agreement; and (iv) protect Confidential Information using the same degree of care '
             'it uses to protect its own confidential information, but in no event less than '
             'reasonable care.')

    add_body(doc, '5.3 Exceptions. The obligations set forth in Section 5.2 shall not apply to '
             'information that: (i) is or becomes publicly available through no fault of the '
             'receiving Party; (ii) was known to the receiving Party prior to disclosure; '
             '(iii) is independently developed by the receiving Party without use of Confidential '
             'Information; or (iv) is required to be disclosed by law, regulation, or court order, '
             'provided that the receiving Party gives prompt notice to the disclosing Party.')

    add_body(doc, '5.4 Duration. The obligations under this Section 5 shall survive the '
             'termination or expiration of this Agreement for a period of five (5) years.')

    # ========== PAGE 5: Intellectual Property and Warranties ==========
    doc.add_page_break()

    add_heading_styled(doc, '6. INTELLECTUAL PROPERTY', level=1)

    add_body(doc, '6.1 Work Product. All deliverables, documentation, source code, and other '
             'materials created by Service Provider in the course of performing the Services '
             '("Work Product") shall be considered works made for hire and shall be the exclusive '
             'property of Client upon creation.')

    add_body(doc, '6.2 Pre-Existing IP. Service Provider retains all right, title, and interest '
             'in its pre-existing intellectual property, tools, frameworks, and methodologies '
             '("Service Provider IP"). To the extent any Service Provider IP is incorporated '
             'into the Work Product, Service Provider hereby grants Client a perpetual, '
             'irrevocable, worldwide, royalty-free license to use such Service Provider IP '
             'solely as embedded in the Work Product.')

    add_body(doc, '6.3 Third-Party Components. Service Provider shall not incorporate any '
             'third-party software or intellectual property into the Work Product without '
             'Client\'s prior written approval. Service Provider shall ensure that all necessary '
             'licenses for third-party components are obtained and properly documented.')

    add_heading_styled(doc, '7. REPRESENTATIONS AND WARRANTIES', level=1)

    add_body(doc, '7.1 Service Provider represents and warrants that: (i) it has the authority '
             'to enter into this Agreement; (ii) the Services will be performed in a professional '
             'and workmanlike manner consistent with generally accepted industry standards; '
             '(iii) the Work Product will not infringe upon any third-party intellectual property '
             'rights; and (iv) all personnel assigned to the Project possess the qualifications '
             'and experience necessary to perform their assigned tasks.')

    add_body(doc, '7.2 Client represents and warrants that: (i) it has the authority to enter '
             'into this Agreement; (ii) it shall provide timely access to systems, data, and '
             'personnel as reasonably required for Service Provider to perform the Services; '
             'and (iii) all information provided to Service Provider in connection with the '
             'Services is accurate and complete.')

    # ========== PAGE 6: Termination, Governing Law, Signatures ==========
    doc.add_page_break()

    add_heading_styled(doc, '8. TERM AND TERMINATION', level=1)

    add_body(doc, '8.1 Term. This Agreement shall commence on the Effective Date and shall '
             'continue until the completion of all Services, unless earlier terminated in '
             'accordance with this Section 8.')

    add_body(doc, '8.2 Termination for Convenience. Either Party may terminate this Agreement '
             'upon sixty (60) days\' prior written notice to the other Party.')

    add_body(doc, '8.3 Termination for Cause. Either Party may terminate this Agreement '
             'immediately upon written notice if the other Party: (i) commits a material breach '
             'that remains uncured for thirty (30) days after written notice; or (ii) becomes '
             'insolvent, files for bankruptcy, or has a receiver appointed for a substantial '
             'portion of its assets.')

    add_heading_styled(doc, '9. GOVERNING LAW AND DISPUTE RESOLUTION', level=1)

    add_body(doc, '9.1 This Agreement shall be governed by and construed in accordance with '
             'the laws of the State of California, without regard to its conflict of laws '
             'principles.')

    add_body(doc, '9.2 Any dispute arising out of or relating to this Agreement shall first '
             'be submitted to non-binding mediation in San Francisco, California. If mediation '
             'fails to resolve the dispute within sixty (60) days, either Party may pursue '
             'binding arbitration under the rules of the American Arbitration Association.')

    add_heading_styled(doc, '10. SIGNATURES', level=1)

    add_body(doc, 'IN WITNESS WHEREOF, the Parties have executed this Agreement as of the '
             'Effective Date.')
    add_body(doc, '')

    # Signature blocks
    add_body(doc, 'MERIDIAN TECHNOLOGIES INC.', bold=True)
    add_body(doc, '')
    add_body(doc, '___________________________________')
    add_body(doc, 'Name: Victoria R. Harrington')
    add_body(doc, 'Title: Chief Operating Officer')
    add_body(doc, 'Date: March 15, 2025')
    add_body(doc, '')
    add_body(doc, '')

    add_body(doc, 'NEXGEN SOLUTIONS GROUP LLC', bold=True)
    add_body(doc, '')
    add_body(doc, '___________________________________')
    add_body(doc, 'Name: David K. Matsumoto')
    add_body(doc, 'Title: Managing Partner')
    add_body(doc, 'Date: March 15, 2025')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
