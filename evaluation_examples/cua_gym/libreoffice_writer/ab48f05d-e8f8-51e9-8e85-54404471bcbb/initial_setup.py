"""
Initial Setup: Navigate tracked changes and accept only insertions
Task ID: writer_rm_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import copy
import datetime

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_043'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_run_element(text, bold=False, italic=False, font_name="Calibri", font_size_pt=11):
    """Create a w:r element with text and optional formatting."""
    r = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr>'
        f'    <w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'    <w:sz w:val="{font_size_pt * 2}"/>'
        f'    <w:szCs w:val="{font_size_pt * 2}"/>'
        + ('    <w:b/>' if bold else '')
        + ('    <w:i/>' if italic else '')
        + f'  </w:rPr>'
        f'  <w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    return r


def make_insertion(rev_id, author, date, text, bold=False, italic=False):
    """Create a w:ins element (tracked insertion)."""
    ins = parse_xml(
        f'<w:ins {nsdecls("w")} w:id="{rev_id}" '
        f'w:author="{author}" w:date="{date}">'
        f'</w:ins>'
    )
    r = make_run_element(text, bold=bold, italic=italic)
    ins.append(r)
    return ins


def make_deletion(rev_id, author, date, text, bold=False, italic=False):
    """Create a w:del element (tracked deletion)."""
    del_elem = parse_xml(
        f'<w:del {nsdecls("w")} w:id="{rev_id}" '
        f'w:author="{author}" w:date="{date}">'
        f'</w:del>'
    )
    # Deletion runs use w:delText instead of w:t
    r = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr>'
        f'    <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
        f'    <w:sz w:val="22"/>'
        f'    <w:szCs w:val="22"/>'
        + ('    <w:b/>' if bold else '')
        + ('    <w:i/>' if italic else '')
        + f'  </w:rPr>'
        f'  <w:delText xml:space="preserve">{text}</w:delText>'
        f'</w:r>'
    )
    del_elem.append(r)
    return del_elem


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Document Title ---
    title = doc.add_heading('Software Release Manual — Version 3.2 Update', level=1)

    # --- Introduction ---
    doc.add_heading('1. Introduction', level=2)
    p1 = doc.add_paragraph()
    p1.add_run(
        'This document outlines the changes, improvements, and known issues '
        'associated with the Version 3.2 release of the Aurora Platform. '
        'All team members should review this manual before deploying the update '
        'to production environments.'
    )

    # --- System Requirements ---
    doc.add_heading('2. System Requirements', level=2)
    p2 = doc.add_paragraph()
    p2.add_run(
        'The Aurora Platform Version 3.2 requires a minimum of 8 GB RAM, '
        '4-core processor, and 50 GB of available disk space. '
    )

    # Now we add tracked changes to paragraph p2
    # Change 1 (INSERTION): Add new requirement text
    # Change 2 (DELETION): Remove old OS requirement

    # --- Installation Procedure ---
    doc.add_heading('3. Installation Procedure', level=2)
    p3 = doc.add_paragraph()
    p3.add_run(
        'Before starting the installation, ensure all running services are stopped '
        'and a full backup of the database has been completed. '
    )

    # --- Configuration Updates ---
    doc.add_heading('4. Configuration Updates', level=2)
    p4 = doc.add_paragraph()
    p4.add_run(
        'Several configuration parameters have been modified in this release. '
        'The primary configuration file is located at /etc/aurora/config.yaml. '
    )

    # --- Security Patches ---
    doc.add_heading('5. Security Patches', level=2)
    p5 = doc.add_paragraph()
    p5.add_run(
        'This release includes critical security updates addressing '
        'vulnerabilities identified in the Q4 2025 security audit. '
    )

    # --- Performance Improvements ---
    doc.add_heading('6. Performance Improvements', level=2)
    p6 = doc.add_paragraph()
    p6.add_run(
        'Query execution time has been reduced by approximately 40 percent '
        'through optimized indexing strategies and connection pooling. '
    )

    # --- Known Issues ---
    doc.add_heading('7. Known Issues', level=2)
    p7 = doc.add_paragraph()
    p7.add_run(
        'The following issues have been identified and are scheduled for '
        'resolution in the next patch release. '
    )

    # --- Rollback Procedure ---
    doc.add_heading('8. Rollback Procedure', level=2)
    p8 = doc.add_paragraph()
    p8.add_run(
        'If critical failures are encountered after deployment, follow the '
        'rollback procedure documented in Appendix B. '
    )

    # Save the base document first
    doc.save(OUTPUT)

    # Now reopen and inject tracked changes via XML manipulation
    doc = Document(OUTPUT)
    body = doc.element.body

    # Get all paragraphs
    paras = body.findall(qn('w:p'))

    # We need to identify the content paragraphs (skip headings).
    # Headings use pStyle with "Heading" prefix.
    content_paras = []
    for p in paras:
        pPr = p.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None and 'Heading' in pStyle.get(qn('w:val'), ''):
                continue
        content_paras.append(p)

    author = "Elena Kowalski"
    base_date = "2026-03-28T10:15:00Z"

    # Tracked changes sequence (10 total):
    # 1. insertion, 2. deletion, 3. insertion, 4. insertion,
    # 5. deletion, 6. insertion, 7. deletion, 8. insertion,
    # 9. insertion, 10. deletion

    # We'll distribute tracked changes across the content paragraphs:
    # Para 0 (Intro): change 1 (ins), change 2 (del)
    # Para 1 (Sys Req): change 3 (ins), change 4 (ins)
    # Para 2 (Install): change 5 (del)
    # Para 3 (Config): change 6 (ins)
    # Para 4 (Security): change 7 (del)
    # Para 5 (Performance): change 8 (ins), change 9 (ins)
    # Para 6 (Known Issues): change 10 (del)

    rev_id = 1

    # Change 1 (INSERTION) - append to para 0 (Introduction)
    ins1 = make_insertion(rev_id, author, "2026-03-28T10:15:00Z",
        "The update also introduces enhanced monitoring capabilities for real-time system health tracking. ")
    content_paras[0].append(ins1)
    rev_id += 1

    # Change 2 (DELETION) - append to para 0
    del2 = make_deletion(rev_id, author, "2026-03-28T10:17:00Z",
        "Legacy compatibility mode has been deprecated in this release. ")
    content_paras[0].append(del2)
    rev_id += 1

    # Change 3 (INSERTION) - append to para 1 (System Requirements)
    ins3 = make_insertion(rev_id, author, "2026-03-28T10:20:00Z",
        "Additionally, TLS 1.3 support is now mandatory for all network communications. ")
    content_paras[1].append(ins3)
    rev_id += 1

    # Change 4 (INSERTION) - append to para 1
    ins4 = make_insertion(rev_id, author, "2026-03-28T10:22:00Z",
        "Docker Engine 24.0 or later is required for containerized deployments. ")
    content_paras[1].append(ins4)
    rev_id += 1

    # Change 5 (DELETION) - append to para 2 (Installation)
    del5 = make_deletion(rev_id, author, "2026-03-28T10:25:00Z",
        "The legacy migration tool should be executed before the main installer. ")
    content_paras[2].append(del5)
    rev_id += 1

    # Change 6 (INSERTION) - append to para 3 (Configuration)
    ins6 = make_insertion(rev_id, author, "2026-03-28T10:30:00Z",
        "A new parameter max_concurrent_sessions has been added with a default value of 250. ")
    content_paras[3].append(ins6)
    rev_id += 1

    # Change 7 (DELETION) - append to para 4 (Security)
    del7 = make_deletion(rev_id, author, "2026-03-28T10:35:00Z",
        "The deprecated SHA-1 certificate validation has been removed from the trust chain. ")
    content_paras[4].append(del7)
    rev_id += 1

    # Change 8 (INSERTION) - append to para 5 (Performance)
    ins8 = make_insertion(rev_id, author, "2026-03-28T10:40:00Z",
        "Memory consumption during batch operations has been reduced by 35 percent. ")
    content_paras[5].append(ins8)
    rev_id += 1

    # Change 9 (INSERTION) - append to para 5
    ins9 = make_insertion(rev_id, author, "2026-03-28T10:42:00Z",
        "Cache invalidation now uses a probabilistic algorithm for improved throughput. ")
    content_paras[5].append(ins9)
    rev_id += 1

    # Change 10 (DELETION) - append to para 6 (Known Issues)
    del10 = make_deletion(rev_id, author, "2026-03-28T10:45:00Z",
        "The automatic failover mechanism may trigger false positives under extreme load conditions. ")
    content_paras[6].append(del10)
    rev_id += 1

    # Save the document with tracked changes
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')
    print(f'Total tracked changes injected: 10 (6 insertions, 4 deletions)')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
