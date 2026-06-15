"""
Initial Setup: Insert an embedded Calc spreadsheet table as an OLE object into a Writer document
Task ID: writer_fp_024
Domain: libreoffice_writer

Creates a 3-page quarterly business report. Page 3 has 'Sales Summary' heading
and 'See table below:' text. No OLE objects exist in the document.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_024'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ===================== PAGE 1: Executive Summary =====================
    heading = doc.add_heading('Quarterly Business Report', level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by: Strategic Planning Division')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Fiscal Year 2025 — Annual Review')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')  # spacer

    doc.add_heading('Executive Summary', level=1)

    doc.add_paragraph(
        'This report provides a comprehensive overview of our company\'s financial '
        'performance during fiscal year 2025. Despite challenging market conditions in '
        'the third quarter, the organization demonstrated resilience and achieved overall '
        'growth targets set at the beginning of the year.'
    )

    doc.add_paragraph(
        'Key highlights include a 5.5% increase in annual revenue, successful expansion '
        'into two new regional markets, and the launch of three innovative product lines '
        'that contributed significantly to our bottom line. The leadership team remains '
        'committed to sustainable growth and operational excellence.'
    )

    doc.add_heading('Market Overview', level=2)

    doc.add_paragraph(
        'The global technology services market experienced moderate growth of 3.2% in 2025, '
        'driven primarily by increased demand for cloud infrastructure and AI-powered '
        'automation solutions. Our positioning in enterprise software and consulting services '
        'allowed us to capture above-average market share gains.'
    )

    doc.add_paragraph(
        'Competition intensified in the mid-market segment, with several new entrants offering '
        'aggressive pricing models. However, our established client relationships and reputation '
        'for quality delivery provided a competitive moat that preserved margins.'
    )

    doc.add_heading('Strategic Initiatives', level=2)

    initiatives = [
        'Digital transformation consulting practice expanded to 14 new enterprise accounts',
        'Cloud migration services grew revenue by 18% year-over-year',
        'Customer satisfaction scores improved from 87% to 92% across all service lines',
        'Employee retention rate reached 94%, well above the industry average of 82%',
        'R&D investment increased to 12% of revenue, funding next-generation platform development',
    ]
    for item in initiatives:
        doc.add_paragraph(item, style='List Bullet')

    # ===================== PAGE BREAK → PAGE 2: Financial Overview =====================
    doc.add_page_break()

    doc.add_heading('Financial Overview', level=1)

    doc.add_paragraph(
        'The financial performance for FY2025 reflects the disciplined execution of our '
        'strategic plan. Total revenue reached $5.3 million, representing a 5.5% increase '
        'from the prior year. Operating expenses were managed effectively, with a cost-to-revenue '
        'ratio improvement of 1.8 percentage points.'
    )

    doc.add_heading('Revenue Breakdown by Division', level=2)

    # Create a revenue table for page 2
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ['Division', 'FY2024 Revenue', 'FY2025 Revenue', 'Growth']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    division_data = [
        ['Enterprise Solutions', '$1,850,000', '$1,960,000', '5.9%'],
        ['Cloud Services', '$1,200,000', '$1,416,000', '18.0%'],
        ['Consulting', '$980,000', '$1,029,000', '5.0%'],
        ['Product Licensing', '$750,000', '$772,500', '3.0%'],
        ['Support & Maintenance', '$245,000', '$257,250', '5.0%'],
    ]
    for r, row_data in enumerate(division_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')  # spacer

    doc.add_heading('Operational Highlights', level=2)

    doc.add_paragraph(
        'Operating margin improved to 18.3% from 16.5% in the prior year, driven by '
        'economies of scale in cloud services and successful automation of internal processes. '
        'The transition to a hybrid work model reduced facility costs by $320,000 annually '
        'without impacting productivity metrics.'
    )

    doc.add_paragraph(
        'Capital expenditure totaled $635,000, primarily allocated to data center upgrades '
        'and security infrastructure modernization. These investments are expected to reduce '
        'operating costs by an estimated $180,000 per year starting in FY2026.'
    )

    doc.add_heading('Workforce Development', level=2)

    doc.add_paragraph(
        'Total headcount grew from 142 to 158 employees, with strategic hires focused on '
        'AI/ML engineering, cybersecurity, and solution architecture roles. The company '
        'invested $420,000 in professional development programs, resulting in 23 employees '
        'obtaining industry-recognized certifications.'
    )

    # ===================== PAGE BREAK → PAGE 3: Sales Summary =====================
    doc.add_page_break()

    doc.add_heading('Sales Summary', level=1)

    doc.add_paragraph(
        'The following section presents the quarterly sales performance for FY2025. '
        'Each quarter reflects the combined revenue from all business divisions, along '
        'with quarter-over-quarter growth percentages. These figures have been audited '
        'and verified by our internal finance team.'
    )

    cursor_para = doc.add_paragraph('See table below:')
    cursor_para.paragraph_format.space_after = Pt(12)

    # Additional content after the placeholder to fill page 3
    doc.add_paragraph('')  # empty space where OLE object would go

    doc.add_paragraph(
        'Note: Revenue figures are presented in millions of US dollars. Growth percentages '
        'are calculated on a quarter-over-quarter basis relative to Q4 of FY2024, which '
        'reported revenue of $1.14 million.'
    )

    doc.add_heading('Regional Performance', level=2)

    doc.add_paragraph(
        'North American operations continued to be the primary revenue driver, accounting '
        'for 62% of total sales. The European division showed strong momentum with 15% '
        'growth, benefiting from new partnerships with major financial institutions in '
        'Frankfurt and London.'
    )

    doc.add_paragraph(
        'The Asia-Pacific region, while still in the early growth phase, exceeded targets '
        'by 8% and secured three landmark deals with telecommunications providers in '
        'Singapore and Sydney. Management expects this region to contribute 20% of total '
        'revenue by FY2027.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
