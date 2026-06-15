"""
FINAL REWARD SCRIPT - SUCCESS
Task: Please update the abstract heading so each word's first letter is uppercase.
Generated: 2025-10-14 10:01:16
Status: success
Model: azure-o3
Total Steps: 9
"""

import os
import re
from docx import Document

def verify_abstract_titlecase(doc_path: str) -> float:
    """Reward script for the task:
    "Please update the abstract heading so each word's first letter is uppercase."

    The script:
    1. Opens the provided DOCX file.
    2. Finds the first paragraph that (a) is styled as a heading and (b) contains the word
       "abstract" (case-insensitive).
    3. Extracts all alphabetic words from that heading and checks whether each word starts
       with an uppercase letter.
    4. Awards a progressive score equal to the proportion of correctly-capitalised words.
       • 1.0  → every word starts with an uppercase letter (perfect completion)
       • <1.0 → partial compliance (some words not capitalised)
       • 0.0  → heading not found or no words correctly capitalised

    The function prints detailed diagnostics and finally prints
        REWARD: <score>
    before returning the score as a float.
    """

    max_score = 1.0

    print(f"Checking file: {doc_path}")

    # 1. Prerequisite: file must exist
    if not os.path.exists(doc_path):
        print("✗ File does not exist – task failed")
        return 0.0

    # 2. Prerequisite: file must load successfully
    try:
        doc = Document(doc_path)
        print(f"✓ Document opened – contains {len(doc.paragraphs)} paragraphs")
    except Exception as e:
        print(f"✗ Failed to load DOCX: {e}")
        return 0.0

    # 3. Locate the abstract heading paragraph (heading style & contains 'abstract')
    abstract_para = None
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.lower().startswith("heading") and "abstract" in para.text.lower():
            abstract_para = para
            break

    if not abstract_para:
        print("✗ No heading paragraph containing the word 'abstract' found")
        return 0.0

    heading_text = abstract_para.text.strip()
    print(f"✓ Located abstract heading: '{heading_text}' (style: {abstract_para.style.name})")

    # 4. Analyse capitalisation of each word
    words = re.findall(r"[A-Za-z']+", heading_text)
    if not words:
        print("✗ Heading contains no alphabetic words – cannot evaluate")
        return 0.0

    correct = 0
    for w in words:
        # Identify first alphabetic character (handles leading quotes/apostrophes)
        first_alpha = next((ch for ch in w if ch.isalpha()), "")
        if first_alpha and first_alpha.isupper():
            correct += 1
        else:
            print(f"  ✗ Word '{w}' does not start with an uppercase letter")

    ratio = correct / len(words)
    print(f"Capitalisation correctness: {correct}/{len(words)} = {ratio:.2f}")

    # 5. Progressive scoring based on correctness ratio
    reward = min(max(ratio * max_score, 0.0), 1.0)

    print(f"REWARD: {reward}")
    return reward


if __name__ == "__main__":
    path = "/home/user/please_update_the_abstract_heading_so_each_words_first_letter_is_uppercase.docx"
    verify_abstract_titlecase(path)
