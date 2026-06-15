"""
Initial Setup: HR policy document with 3 paragraphs, all using default line spacing.
Task ID: osworld_writer_line_spacing_per_paragraph_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_line_spacing_per_paragraph_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Introduction / Purpose section
    # All paragraphs use default (inherited) line spacing — no custom spacing set
    p1 = doc.add_paragraph(
        "This Employee Remote Work Policy establishes guidelines and expectations for "
        "employees who work from locations other than the company's primary office. "
        "The policy applies to all full-time and part-time employees who have received "
        "written approval from their department manager to work remotely on a regular "
        "or occasional basis. Compliance with this policy is mandatory for all remote "
        "work arrangements."
    )
    # Do NOT set any line_spacing — leave at default (inherited)

    # Paragraph 2: Eligibility section
    p2 = doc.add_paragraph(
        "Eligibility for remote work is determined based on job responsibilities, "
        "performance history, and operational requirements. Employees must have "
        "completed a minimum of six months of continuous employment and maintained "
        "a satisfactory performance rating in their most recent annual review. "
        "Positions that require regular in-person collaboration, hands-on lab work, "
        "or direct client interaction may not be eligible for remote arrangements. "
        "Final approval rests with the department head and Human Resources."
    )
    # Do NOT set any line_spacing — leave at default (inherited)

    # Paragraph 3: Responsibilities section
    p3 = doc.add_paragraph(
        "Remote employees are expected to maintain their standard working hours and "
        "remain accessible via company communication tools during their scheduled "
        "shifts. They must ensure a dedicated, distraction-free workspace that meets "
        "ergonomic and security standards. All company data must be accessed through "
        "the approved VPN connection, and employees are prohibited from using public "
        "or unsecured networks for work-related activities. Equipment provided by the "
        "company remains the property of the organization and must be returned upon "
        "termination of employment or the remote work arrangement."
    )
    # Do NOT set any line_spacing — leave at default (inherited)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
