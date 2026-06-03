"""
Reward Script: Set different line spacing for the first two paragraphs
Task ID: osworld_writer_line_spacing_per_paragraph_003
Domain: libreoffice_writer
Scoring:
  Component 1: Paragraph 1 has single line spacing (1.0)  — 0.4 pts
  Component 2: Paragraph 2 has double line spacing (2.0)  — 0.4 pts
  Component 3: Paragraph 3 remains at default spacing     — 0.2 pts
  Total: 1.0
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_line_spacing_per_paragraph_003'


def verify_task(file_path):
    """
    Verify that:
      - Paragraph 1 (index 0) has single line spacing set to 1.0
      - Paragraph 2 (index 1) has double line spacing set to 2.0
      - Paragraph 3 (index 2) is unchanged (default/None spacing)

    Returns a float between 0.0 and 1.0.
    """
    total_score = 0.0

    # Load the document — fail fast if not readable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 3 paragraphs
    if len(doc.paragraphs) < 3:
        print(f"CRITICAL: Expected at least 3 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    para0 = doc.paragraphs[0]
    para1 = doc.paragraphs[1]
    para2 = doc.paragraphs[2]

    # Component 1: Paragraph 1 line spacing == 1.0 (single) (0.4 pts)
    # This FAILS on initial_env (None) and PASSES on golden_env (1.0)
    try:
        ls0 = para0.paragraph_format.line_spacing
        # Accept both exact float 1.0 and the round-trip from python-docx SINGLE rule
        if ls0 is not None and abs(float(ls0) - 1.0) < 0.05:
            print(f"PASS: Component 1 — paragraph 1 has single (1.0) line spacing (found {ls0}) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — expected paragraph 1 line_spacing=1.0, found {ls0}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Paragraph 2 line spacing == 2.0 (double) (0.4 pts)
    # This FAILS on initial_env (None) and PASSES on golden_env (2.0)
    try:
        ls1 = para1.paragraph_format.line_spacing
        if ls1 is not None and abs(float(ls1) - 2.0) < 0.05:
            print(f"PASS: Component 2 — paragraph 2 has double (2.0) line spacing (found {ls1}) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 — expected paragraph 2 line_spacing=2.0, found {ls1}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Paragraph 3 is unchanged — line spacing remains None (default) (0.2 pts)
    # In both initial and golden, para 3 should be None.
    # This check prevents regressions where all paragraphs are blindly modified.
    # It PASSES on initial_env when para3 is also None, BUT because Components 1 and 2
    # will FAIL on initial_env, the overall initial score will still be 0.0.
    # We only award these points if Components 1+2 passed (to avoid partial credit on initial).
    try:
        ls2 = para2.paragraph_format.line_spacing
        # Paragraph 3 must remain at default (None) — no explicit spacing set
        if ls2 is None and total_score >= 0.8:
            # Only award if C1 and C2 also passed, ensuring this is not a false positive
            print(f"PASS: Component 3 — paragraph 3 remains at default line spacing (None) (0.2 pts)")
            total_score += 0.2
        elif ls2 is not None:
            print(f"FAIL: Component 3 — paragraph 3 should have default spacing, found {ls2}")
        else:
            # ls2 is None but C1/C2 failed — don't award points
            print(f"SKIP: Component 3 — paragraph 3 spacing is None but preceding components not satisfied")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM environment
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
