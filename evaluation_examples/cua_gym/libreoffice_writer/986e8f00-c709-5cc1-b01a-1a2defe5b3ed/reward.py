"""
Reward Script: Format customer success story as downloadable PDF guide
Task ID: writer_mktg_053
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Header paragraph has dark blue (#0A2463) background shading and white bold text
  Component 2 (0.20): Title paragraph formatted as 22pt bold with dark blue (#0A2463) color
  Component 3 (0.25): Section headings (Challenge/Solution/Results) have bottom border dividers
  Component 4 (0.30): Customer quote formatted as pull quote (decorative quote mark + 18pt italic centered + gray background)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'customer_success_story'
FILE_PATH = f'{WORKDIR}/{TASK_ID}.docx'

# XML namespace attribute helpers
W_VAL   = qn('w:val')
W_COLOR = qn('w:color')
W_SZ    = qn('w:sz')
W_FILL  = qn('w:fill')


def color_close(hex_str, target_hex, tolerance=30):
    """Check if a hex color string is close to target (simple RGB distance)."""
    if hex_str is None:
        return False
    try:
        r1 = int(hex_str[0:2], 16)
        g1 = int(hex_str[2:4], 16)
        b1 = int(hex_str[4:6], 16)
        r2 = int(target_hex[0:2], 16)
        g2 = int(target_hex[2:4], 16)
        b2 = int(target_hex[4:6], 16)
        dist = ((r1-r2)**2 + (g1-g2)**2 + (b1-b2)**2) ** 0.5
        return dist <= tolerance
    except Exception:
        return False


def get_para_shading_fill(para):
    """Return the fill hex string of paragraph shading, or None."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        return None
    return shd.get(W_FILL)


def get_para_bottom_border(para):
    """Return (val, color, sz) tuple for paragraph bottom border, or None."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        return None
    b = pBdr.find(qn('w:bottom'))
    if b is None:
        return None
    return (b.get(W_VAL), b.get(W_COLOR), b.get(W_SZ))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: Header has dark blue (#0A2463) background with white text (0.25 pts)
    # The header paragraph shading fill must be ~#0A2463 and text must be white+bold.
    # -------------------------------------------------------------------------
    try:
        section = doc.sections[0]
        header = section.header
        header_pass = False
        header_details = "No header paragraphs found"

        for p in header.paragraphs:
            fill = get_para_shading_fill(p)
            # Check background color is dark blue
            bg_ok = color_close(fill, '0A2463', tolerance=30) if fill else False

            # Check text color is white and bold
            white_bold_ok = False
            for run in p.runs:
                if run.text.strip():
                    try:
                        rc = run.font.color.rgb
                        # White: FFFFFF or very close
                        color_ok = color_close(str(rc), 'FFFFFF', tolerance=30) if rc else False
                        bold_ok = run.bold is True
                        if color_ok and bold_ok:
                            white_bold_ok = True
                            break
                    except Exception:
                        pass

            if bg_ok and white_bold_ok:
                header_pass = True
                header_details = f"fill={fill}, white+bold text confirmed"
                break
            elif bg_ok:
                header_details = f"fill={fill} (bg OK, but white+bold text not confirmed)"
            elif white_bold_ok:
                header_details = f"white+bold text found but bg fill={fill} not dark blue"

        if header_pass:
            print(f"PASS: Component 1 — Header has dark blue background with white bold text ({header_details}) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Header missing dark blue bg or white bold text: {header_details}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Title paragraph is 22pt bold with dark blue (#0A2463) color (0.20 pts)
    # The first paragraph (title) must have font size >= 20pt, bold, and dark blue color.
    # -------------------------------------------------------------------------
    try:
        title_keywords = ['From Manual Chaos', 'Meridian Health Transformed']
        title_para = None
        for para in doc.paragraphs:
            if any(kw in para.text for kw in title_keywords):
                title_para = para
                break

        if title_para is None:
            print("FAIL: Component 2 — Title paragraph not found")
        else:
            size_ok = False
            bold_ok = False
            color_ok = False
            details = []
            for run in title_para.runs:
                if not run.text.strip():
                    continue
                # Size check: >= 20pt (ground truth says 22pt)
                if run.font.size and run.font.size.pt >= 20:
                    size_ok = True
                    details.append(f"size={run.font.size.pt}pt")
                # Bold check
                if run.bold is True:
                    bold_ok = True
                    details.append("bold=True")
                # Color check: dark blue #0A2463
                try:
                    rc = run.font.color.rgb
                    if rc and color_close(str(rc), '0A2463', tolerance=40):
                        color_ok = True
                        details.append(f"color={rc}")
                except Exception:
                    pass

            if size_ok and bold_ok and color_ok:
                print(f"PASS: Component 2 — Title is 22pt bold dark blue ({', '.join(details)}) (0.20 pts)")
                total_score += 0.20
            else:
                missing = []
                if not size_ok:
                    missing.append("size<20pt")
                if not bold_ok:
                    missing.append("not bold")
                if not color_ok:
                    missing.append("color not dark blue")
                print(f"FAIL: Component 2 — Title missing: {missing}; details: {details}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Section headings have bottom border dividers (0.25 pts)
    # At least 2 of the 3 main sections (Challenge/Solution/Results) must have
    # a bottom border with color ~#0A2463.
    # -------------------------------------------------------------------------
    try:
        section_headings = ['The Challenge', 'The Solution', 'The Results']
        borders_found = 0
        borders_detail = []

        for para in doc.paragraphs:
            if para.text.strip() in section_headings:
                border_info = get_para_bottom_border(para)
                if border_info:
                    val, bcolor, sz = border_info
                    if val and val != 'none' and bcolor and color_close(bcolor, '0A2463', tolerance=40):
                        borders_found += 1
                        borders_detail.append(f"{para.text.strip()}: val={val}, color={bcolor}, sz={sz}")

        if borders_found >= 2:
            print(f"PASS: Component 3 — {borders_found}/3 section headings have bottom borders ({'; '.join(borders_detail)}) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 3 — Only {borders_found}/3 section headings have bottom borders (need >=2); found: {borders_detail}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: Customer quote formatted as pull quote (0.30 pts)
    # Checks:
    #   (a) A decorative open-quote character paragraph exists (large font >=36pt) - 0.10 pts
    #   (b) Quote text is 18pt italic and centered - 0.10 pts
    #   (c) Quote area has light gray (#F5F5F5) background shading - 0.10 pts
    # -------------------------------------------------------------------------
    try:
        quote_text_fragment = "Apex Dynamics didn"  # fragment of the actual quote
        attribution_fragment = "Sarah Mitchell"

        # Find quote-related paragraphs
        decorative_quote_para = None
        quote_text_para = None
        attribution_para = None
        gray_background_count = 0

        for para in doc.paragraphs:
            # Look for decorative quote mark paragraph (large font)
            for run in para.runs:
                if run.text.strip() in ['\u201c', '"', '\u2018', "'", '"', '"']:
                    if run.font.size and run.font.size.pt >= 36:
                        decorative_quote_para = para

            # Look for actual quote text
            if quote_text_fragment in para.text:
                quote_text_para = para

            # Look for attribution
            if attribution_fragment in para.text:
                attribution_para = para

        # Sub-check (a): Decorative quote mark exists
        deco_ok = decorative_quote_para is not None
        if deco_ok:
            run_size = None
            for r in decorative_quote_para.runs:
                if r.font.size:
                    run_size = r.font.size.pt
            print(f"  Sub-check (a): Decorative quote mark found at size={run_size}pt")
        else:
            print(f"  Sub-check (a): Decorative quote mark NOT found (need large-font " or " character)")

        # Sub-check (b): Quote text is 18pt italic and centered
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        quote_italic_ok = False
        quote_size_ok = False
        quote_centered_ok = False
        if quote_text_para:
            align = quote_text_para.paragraph_format.alignment
            if align in [WD_PARAGRAPH_ALIGNMENT.CENTER, 1]:
                quote_centered_ok = True
            for run in quote_text_para.runs:
                if run.text.strip():
                    if run.italic is True:
                        quote_italic_ok = True
                    if run.font.size and run.font.size.pt >= 16:
                        quote_size_ok = True
            print(f"  Sub-check (b): italic={quote_italic_ok}, size_ok={quote_size_ok}, centered={quote_centered_ok}")
        else:
            print(f"  Sub-check (b): Quote text paragraph not found")

        # Sub-check (c): Gray background on quote paragraphs
        gray_paras = []
        for para in doc.paragraphs:
            if quote_text_fragment in para.text or attribution_fragment in para.text:
                fill = get_para_shading_fill(para)
                if fill and color_close(fill, 'F5F5F5', tolerance=40):
                    gray_paras.append(para.text[:30])
        gray_ok = len(gray_paras) >= 1
        print(f"  Sub-check (c): gray background on {len(gray_paras)} quote paragraph(s): {gray_paras}")

        # Score Component 4
        component4_score = 0.0
        if deco_ok:
            component4_score += 0.10
        if quote_italic_ok and quote_size_ok and quote_centered_ok:
            component4_score += 0.10
        if gray_ok:
            component4_score += 0.10

        if component4_score > 0:
            print(f"PASS: Component 4 — Pull quote formatting: {component4_score}/0.30 pts")
            total_score += component4_score
        else:
            print(f"FAIL: Component 4 — Pull quote not formatted correctly (score: 0.0/0.30)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
