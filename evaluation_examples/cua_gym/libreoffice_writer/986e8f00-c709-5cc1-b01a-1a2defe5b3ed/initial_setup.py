"""
Initial Setup: Customer success story document - plain minimal formatting
Task ID: writer_mktg_053
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_053'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/customer_success_story.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set default font for the document body
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)

    # --- Title ---
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(
        'From Manual Chaos to Automated Excellence: How Meridian Health Transformed Patient Engagement'
    )
    title_run.font.size = Pt(12)
    title_run.font.bold = False

    doc.add_paragraph()  # spacer

    # --- Introduction paragraph ---
    intro_para = doc.add_paragraph()
    intro_run = intro_para.add_run(
        'Meridian Health is a regional healthcare network serving over 200,000 patients across '
        'three counties in the Pacific Northwest. In 2023, facing growing administrative complexity '
        'and declining patient satisfaction scores, they partnered with Apex Dynamics to overhaul '
        'their patient engagement operations.'
    )
    intro_run.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # --- The Challenge section ---
    challenge_heading = doc.add_paragraph()
    challenge_run = challenge_heading.add_run('The Challenge')
    challenge_run.font.size = Pt(12)
    challenge_run.font.bold = False

    doc.add_paragraph()  # spacer

    challenge_body1 = doc.add_paragraph()
    challenge_body1.add_run(
        'Meridian Health\'s patient engagement team was overwhelmed. Appointment reminders were '
        'sent manually via spreadsheets, follow-up calls were tracked in disconnected systems, '
        'and patient feedback went largely uncollected. Staff spent over 40% of their time on '
        'administrative tasks instead of direct patient support.'
    ).font.size = Pt(12)

    challenge_body2 = doc.add_paragraph()
    challenge_body2.add_run(
        'The consequences were significant: a 23% no-show rate for appointments, a patient '
        'satisfaction score of 62 out of 100, and staff burnout leading to 18% annual turnover '
        'in the engagement department. Leadership recognized that manual processes were '
        'unsustainable as the network prepared to add two new facilities.'
    ).font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # --- The Solution section ---
    solution_heading = doc.add_paragraph()
    solution_run = solution_heading.add_run('The Solution')
    solution_run.font.size = Pt(12)
    solution_run.font.bold = False

    doc.add_paragraph()  # spacer

    solution_body1 = doc.add_paragraph()
    solution_body1.add_run(
        'Apex Dynamics deployed their EngageCare platform — an AI-powered patient communication '
        'suite — across all Meridian Health facilities over a 90-day implementation window. The '
        'platform integrated with Meridian\'s existing Epic EHR system and automated appointment '
        'reminders, post-visit follow-ups, and satisfaction surveys through patients\' preferred '
        'communication channels.'
    ).font.size = Pt(12)

    solution_body2 = doc.add_paragraph()
    solution_body2.add_run(
        'Key capabilities included: intelligent scheduling optimization that predicted no-show '
        'risk and sent targeted reminders; automated care pathway nudges for chronic condition '
        'management; real-time satisfaction feedback loops with escalation protocols for at-risk '
        'patients; and a unified dashboard giving staff complete visibility into patient journey status.'
    ).font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # --- The Results section ---
    results_heading = doc.add_paragraph()
    results_run = results_heading.add_run('The Results')
    results_run.font.size = Pt(12)
    results_run.font.bold = False

    doc.add_paragraph()  # spacer

    results_body1 = doc.add_paragraph()
    results_body1.add_run(
        'Within six months of full deployment, Meridian Health saw transformative improvements '
        'across every key metric. The no-show rate dropped from 23% to 8%, a 65% reduction that '
        'freed up over 1,200 appointment slots per month. Patient satisfaction scores climbed from '
        '62 to 89 out of 100, placing Meridian in the top quartile of regional health networks.'
    ).font.size = Pt(12)

    results_body2 = doc.add_paragraph()
    results_body2.add_run(
        'Staff time spent on administrative tasks fell to 15%, allowing the engagement team to '
        'redirect their energy toward high-value patient interactions. Annual staff turnover in '
        'the department dropped to 7%. Revenue capture improved by $2.3 million annually due to '
        'reduced no-shows and better follow-through on recommended care pathways.'
    ).font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # --- Customer Quote section ---
    quote_heading = doc.add_paragraph()
    quote_run = quote_heading.add_run('Customer Quote')
    quote_run.font.size = Pt(12)
    quote_run.font.bold = False

    doc.add_paragraph()  # spacer

    # Quote paragraph — plain, no pull-quote styling, no decorative elements
    quote_para = doc.add_paragraph()
    quote_run2 = quote_para.add_run(
        '\u201cApex Dynamics didn\u2019t just solve our operational problems \u2014 they transformed '
        'how we think about patient engagement entirely.\u201d \u2014 Sarah Mitchell, COO, Meridian Health'
    )
    quote_run2.font.size = Pt(12)
    quote_run2.font.italic = False

    doc.add_paragraph()  # spacer

    # --- About Meridian Health section ---
    about_heading = doc.add_paragraph()
    about_run = about_heading.add_run('About Meridian Health')
    about_run.font.size = Pt(12)
    about_run.font.bold = False

    doc.add_paragraph()  # spacer

    about_body = doc.add_paragraph()
    about_body.add_run(
        'Meridian Health is a not-for-profit regional healthcare network headquartered in Portland, '
        'Oregon. Founded in 1987, the network operates four acute care hospitals, twelve outpatient '
        'clinics, and a robust telehealth program. Meridian Health is accredited by The Joint '
        'Commission and consistently recognized for clinical excellence and community health '
        'initiatives. For more information, visit www.meridianhealth.org.'
    ).font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
