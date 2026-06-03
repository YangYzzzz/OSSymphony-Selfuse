"""
Reward Script: Add a cover page to an employee handbook
Task ID: writer_hr_051
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): "Summit Financial Group" present, centered, 24pt, bold
  Component 2 (0.25): "Employee Handbook" present, centered, 20pt
  Component 3 (0.25): "Effective Date: January 1, 2026" and "Version 3.0" present, centered
  Component 4 (0.25): Page break between cover page and TOC content
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_051'


def persist_app_state(domain):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        import time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def find_para_with_text(doc, target_text, search_range=None):
    """Find a paragraph containing target_text (case-insensitive).
    Returns (index, paragraph) or (None, None)."""
    paras = doc.paragraphs if search_range is None else doc.paragraphs[:search_range]
    for i, para in enumerate(paras):
        if target_text.lower() in para.text.lower().strip():
            return i, para
    return None, None


def is_centered(para):
    """Check if paragraph alignment is CENTER."""
    return para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER


def check_font_size_pt(para, expected_pt, tolerance=1.0):
    """Check if any run in the paragraph has the expected font size (within tolerance)."""
    for run in para.runs:
        if run.font.size is not None:
            actual_pt = run.font.size.pt
            if abs(actual_pt - expected_pt) <= tolerance:
                return True, actual_pt
    return False, None


def check_bold(para):
    """Check if any run with text in the paragraph is bold."""
    for run in para.runs:
        if run.text.strip() and run.font.bold:
            return True
    return False


def has_page_break_before_toc(doc):
    """Check if there is a page break between the cover page content and the TOC.
    Looks for a page break (w:br type=page) in paragraphs that appear before 'Table of Contents'."""
    toc_idx = None
    for i, para in enumerate(doc.paragraphs):
        if 'table of contents' in para.text.lower().strip():
            toc_idx = i
            break

    if toc_idx is None:
        return False, "TOC not found"

    # Look for page break in paragraphs before the TOC
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for i in range(toc_idx):
        para = doc.paragraphs[i]
        for run in para.runs:
            for br in run.element.findall(f'{{{ns_w}}}br'):
                br_type = br.attrib.get(f'{{{ns_w}}}type', 'line')
                if br_type == 'page':
                    return True, f"Page break found at para {i}"

    # Also check if the TOC paragraph itself has page_break_before
    toc_para = doc.paragraphs[toc_idx]
    if toc_para.paragraph_format.page_break_before:
        return True, f"TOC para has page_break_before=True"

    return False, "No page break found before TOC"


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have enough paragraphs for a cover page + content
    if len(doc.paragraphs) < 5:
        print(f"FAIL: Document too short ({len(doc.paragraphs)} paragraphs) - no cover page possible")
        print("REWARD: 0.0")
        return 0.0

    # Find TOC to determine where cover page ends
    toc_idx = None
    for i, para in enumerate(doc.paragraphs):
        if 'table of contents' in para.text.lower().strip():
            toc_idx = i
            break

    if toc_idx is None:
        print("FAIL: 'Table of Contents' not found - document structure may be corrupted")
        print("REWARD: 0.0")
        return 0.0

    # Cover page is everything before the TOC
    cover_paras = doc.paragraphs[:toc_idx]

    # Component 1: "Summit Financial Group" in 24pt bold centered (0.25 points)
    try:
        comp1_idx, comp1_para = None, None
        for i, para in enumerate(cover_paras):
            if 'summit financial group' in para.text.lower().strip():
                comp1_idx, comp1_para = i, para
                break

        if comp1_para is not None:
            is_center = is_centered(comp1_para)
            size_ok, actual_size = check_font_size_pt(comp1_para, 24.0, tolerance=1.0)
            is_bold = check_bold(comp1_para)

            if is_center and size_ok and is_bold:
                print(f"PASS: Component 1 — 'Summit Financial Group' at para {comp1_idx}: centered, {actual_size}pt, bold (0.25 pts)")
                total_score += 0.25
            else:
                reasons = []
                if not is_center:
                    reasons.append(f"not centered (align={comp1_para.paragraph_format.alignment})")
                if not size_ok:
                    reasons.append(f"size not 24pt (actual={actual_size})")
                if not is_bold:
                    reasons.append("not bold")
                print(f"FAIL: Component 1 — 'Summit Financial Group' found but: {', '.join(reasons)}")
        else:
            print("FAIL: Component 1 — 'Summit Financial Group' not found in cover page area")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: "Employee Handbook" in 20pt centered (0.25 points)
    try:
        comp2_idx, comp2_para = None, None
        for i, para in enumerate(cover_paras):
            if 'employee handbook' in para.text.lower().strip():
                comp2_idx, comp2_para = i, para
                break

        if comp2_para is not None:
            is_center = is_centered(comp2_para)
            size_ok, actual_size = check_font_size_pt(comp2_para, 20.0, tolerance=1.0)

            # Verify it appears after company name
            after_company = comp1_idx is not None and comp2_idx > comp1_idx

            if is_center and size_ok:
                print(f"PASS: Component 2 — 'Employee Handbook' at para {comp2_idx}: centered, {actual_size}pt (0.25 pts)")
                total_score += 0.25
            else:
                reasons = []
                if not is_center:
                    reasons.append(f"not centered (align={comp2_para.paragraph_format.alignment})")
                if not size_ok:
                    reasons.append(f"size not 20pt (actual={actual_size})")
                print(f"FAIL: Component 2 — 'Employee Handbook' found but: {', '.join(reasons)}")
        else:
            print("FAIL: Component 2 — 'Employee Handbook' not found in cover page area")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: "Effective Date: January 1, 2026" and "Version 3.0" present and centered (0.25 points)
    try:
        date_para = None
        version_para = None

        for i, para in enumerate(cover_paras):
            text = para.text.strip().lower()
            if 'effective date' in text and '2026' in text:
                date_para = para
            if 'version 3.0' in text:
                version_para = para

        sub_score = 0.0
        if date_para is not None and is_centered(date_para):
            sub_score += 0.125
            print(f"PASS: Component 3a — 'Effective Date: January 1, 2026' found and centered")
        elif date_para is not None:
            sub_score += 0.0625
            print(f"PARTIAL: Component 3a — 'Effective Date' found but not centered")
        else:
            print("FAIL: Component 3a — 'Effective Date: January 1, 2026' not found in cover page")

        if version_para is not None and is_centered(version_para):
            sub_score += 0.125
            print(f"PASS: Component 3b — 'Version 3.0' found and centered")
        elif version_para is not None:
            sub_score += 0.0625
            print(f"PARTIAL: Component 3b — 'Version 3.0' found but not centered")
        else:
            print("FAIL: Component 3b — 'Version 3.0' not found in cover page")

        if sub_score > 0:
            print(f"PASS: Component 3 — date/version ({sub_score} pts)")
            total_score += sub_score
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Page break between cover page and TOC (0.25 points)
    try:
        pb_found, pb_detail = has_page_break_before_toc(doc)
        if pb_found:
            print(f"PASS: Component 4 — Page break before TOC: {pb_detail} (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 4 — {pb_detail}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
