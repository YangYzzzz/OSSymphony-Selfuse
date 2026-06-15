"""
Reward Script: Set up document for professional book printing with A5 page size,
mirrored margins, header/footer, and different first page.
Task ID: writer_page_054
Domain: libreoffice_writer

Scoring Rubric:
  Component 1: Page size is A5 (14.8cm x 21.0cm)         — 0.25 pts
  Component 2: Mirrored margins enabled with correct values — 0.30 pts
               (inner=2.5cm, outer=1.5cm, top=2.0cm, bottom=2.0cm)
  Component 3: Header contains 'The Art of Coding' centered — 0.20 pts
  Component 4: Footer has centered PAGE number field code  — 0.15 pts
  Component 5: Different first page header/footer enabled   — 0.10 pts
  Total: 1.00
"""

import os

from docx import Document
from docx.shared import Cm

WORKDIR = '/home/user/Desktop'
TASK_ID = 'book_manuscript'
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Tolerance: 0.1 cm to account for twip rounding
MARGIN_TOLERANCE_CM = 0.1


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if len(doc.sections) == 0:
        print("CRITICAL: Document has no sections.")
        print("REWARD: 0.0")
        return 0.0

    s = doc.sections[0]

    # -----------------------------------------------------------------
    # Component 1: Page size is A5 (14.8 cm x 21.0 cm)  (0.25 points)
    # -----------------------------------------------------------------
    # A5 = 14.8 cm wide x 21.0 cm tall
    # Initial was A4 = 21.0 cm wide x 29.7 cm tall
    try:
        w_cm = s.page_width.cm
        h_cm = s.page_height.cm
        a5_w = 14.8
        a5_h = 21.0
        w_ok = abs(w_cm - a5_w) < 0.15
        h_ok = abs(h_cm - a5_h) < 0.15
        if w_ok and h_ok:
            print(f"PASS: Component 1 — A5 page size ({w_cm:.4f} x {h_cm:.4f} cm) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Expected A5 (14.8x21.0 cm), found ({w_cm:.4f}x{h_cm:.4f} cm)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------
    # Component 2: Mirrored margins with correct values (0.30 points)
    # mirrorMargins in document settings + inner=2.5cm, outer=1.5cm,
    # top=2.0cm, bottom=2.0cm
    # In mirror margin mode: left margin = inner (odd pages), right = outer (even pages)
    # -----------------------------------------------------------------
    try:
        # Check mirrorMargins is set in document settings
        settings_elem = doc.settings.element
        mirror_elem = settings_elem.find('{%s}mirrorMargins' % WNS)
        mirror_enabled = mirror_elem is not None

        # Check margin values
        left_cm = s.left_margin.cm   # inner margin
        right_cm = s.right_margin.cm  # outer margin
        top_cm = s.top_margin.cm
        bottom_cm = s.bottom_margin.cm

        inner_ok = abs(left_cm - 2.5) < MARGIN_TOLERANCE_CM
        outer_ok = abs(right_cm - 1.5) < MARGIN_TOLERANCE_CM
        top_ok = abs(top_cm - 2.0) < MARGIN_TOLERANCE_CM
        bottom_ok = abs(bottom_cm - 2.0) < MARGIN_TOLERANCE_CM

        margins_ok = inner_ok and outer_ok and top_ok and bottom_ok

        if mirror_enabled and margins_ok:
            print(f"PASS: Component 2 — Mirror margins enabled, inner={left_cm:.4f}cm, "
                  f"outer={right_cm:.4f}cm, top={top_cm:.4f}cm, bottom={bottom_cm:.4f}cm (0.30 pts)")
            total_score += 0.30
        else:
            reasons = []
            if not mirror_enabled:
                reasons.append("mirrorMargins not enabled in document settings")
            if not inner_ok:
                reasons.append(f"inner margin={left_cm:.4f}cm (expected ~2.5cm)")
            if not outer_ok:
                reasons.append(f"outer margin={right_cm:.4f}cm (expected ~1.5cm)")
            if not top_ok:
                reasons.append(f"top margin={top_cm:.4f}cm (expected ~2.0cm)")
            if not bottom_ok:
                reasons.append(f"bottom margin={bottom_cm:.4f}cm (expected ~2.0cm)")
            print(f"FAIL: Component 2 — {'; '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------
    # Component 3: Header contains 'The Art of Coding' centered (0.20 points)
    # Initial: empty header, not linked out
    # Golden: header text = 'The Art of Coding', alignment = CENTER
    # -----------------------------------------------------------------
    try:
        header = s.header
        header_text = ""
        header_alignment = None
        for p in header.paragraphs:
            if p.text.strip():
                header_text = p.text.strip()
                header_alignment = p.alignment
                break

        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        text_ok = "The Art of Coding" in header_text
        align_ok = header_alignment == WD_PARAGRAPH_ALIGNMENT.CENTER

        if text_ok and align_ok:
            print(f"PASS: Component 3 — Header contains '{header_text}' centered (0.20 pts)")
            total_score += 0.20
        else:
            reasons = []
            if not text_ok:
                reasons.append(f"header text={repr(header_text)} (expected 'The Art of Coding')")
            if not align_ok:
                reasons.append(f"header alignment={header_alignment} (expected CENTER)")
            print(f"FAIL: Component 3 — {'; '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------
    # Component 4: Footer has centered PAGE number field code (0.15 points)
    # Initial: empty footer
    # Golden: footer has ' PAGE ' instrText, alignment = CENTER
    # -----------------------------------------------------------------
    try:
        footer = s.footer
        footer_has_page_field = False
        footer_centered = False

        for p in footer.paragraphs:
            # Check for PAGE field code
            instrs = p._p.findall('.//{%s}instrText' % WNS)
            for instr in instrs:
                if instr.text and 'PAGE' in instr.text.upper():
                    footer_has_page_field = True
                    break

            # Check alignment (even if text is empty, field codes set alignment)
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            if p.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                footer_centered = True

        if footer_has_page_field and footer_centered:
            print(f"PASS: Component 4 — Footer has centered PAGE field code (0.15 pts)")
            total_score += 0.15
        else:
            reasons = []
            if not footer_has_page_field:
                reasons.append("no PAGE field code found in footer")
            if not footer_centered:
                reasons.append("footer paragraph not centered")
            print(f"FAIL: Component 4 — {'; '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -----------------------------------------------------------------
    # Component 5: Different first page enabled (titlePg set) (0.10 points)
    # Initial: titlePg absent, different_first_page_header_footer = False
    # Golden: titlePg present, different_first_page_header_footer = True,
    #         first page header/footer are blank
    # -----------------------------------------------------------------
    try:
        diff_first = s.different_first_page_header_footer

        # Verify first page header is blank (empty or whitespace only)
        first_header_blank = True
        if diff_first:
            try:
                first_hdr = s.first_page_header
                for p in first_hdr.paragraphs:
                    if p.text.strip():
                        first_header_blank = False
                        break
            except Exception:
                first_header_blank = True  # if accessor fails, can't verify

        if diff_first:
            print(f"PASS: Component 5 — Different first page header/footer enabled, "
                  f"first header blank={first_header_blank} (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — different_first_page_header_footer={diff_first} (expected True)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -----------------------------------------------------------------
    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
