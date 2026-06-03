"""
Initial Setup: Book manuscript for page setup task
Task ID: writer_page_054
Domain: libreoffice_writer

Creates a 25-page book manuscript on ~/Desktop/ with:
- Page size A4, portrait
- Regular margins: top=2.54cm, bottom=2.54cm, left=2.54cm, right=2.54cm
- No header, no footer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'book_manuscript'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Configure A4 page size with standard margins (no header/footer)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Ensure no header/footer is linked (leave default empty headers/footers)
    # By default python-docx creates empty headers/footers; ensure they have no content
    # and are not enabled (default behavior)

    # ----------------------------------------------------------------
    # Book: "The Art of Coding"
    # ----------------------------------------------------------------

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run('The Art of Coding')
    run.bold = True
    run.font.size = Pt(28)

    doc.add_paragraph()

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle_para.add_run('A Journey Through Software Craftsmanship')
    run.font.size = Pt(16)
    run.italic = True

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author_para.add_run('By Jonathan M. Reeves')
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    publisher_para = doc.add_paragraph()
    publisher_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = publisher_para.add_run('Northbridge Publishing\nSan Francisco, CA\n2024')
    run.font.size = Pt(11)

    # Page break after title page
    doc.add_page_break()

    # Table of Contents page
    toc_heading = doc.add_paragraph('Table of Contents')
    toc_heading.runs[0].bold = True
    toc_heading.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    chapters = [
        ('Chapter 1', 'The Foundation of Clean Code', '7'),
        ('Chapter 2', 'Naming Things Well', '19'),
        ('Chapter 3', 'Functions and Their Secrets', '31'),
        ('Chapter 4', 'Comments: When and Why', '45'),
        ('Chapter 5', 'Error Handling Done Right', '57'),
        ('Chapter 6', 'Unit Testing Philosophy', '69'),
        ('Chapter 7', 'The Art of Refactoring', '83'),
        ('Chapter 8', 'Design Patterns in Practice', '97'),
        ('Chapter 9', 'Performance and Optimization', '113'),
        ('Chapter 10', 'The Human Side of Coding', '127'),
    ]

    for chapter_num, chapter_title, page_num in chapters:
        toc_line = doc.add_paragraph()
        run = toc_line.add_run(f'{chapter_num}: {chapter_title}')
        run.font.size = Pt(11)
        toc_line.paragraph_format.space_after = Pt(4)

    # Page break
    doc.add_page_break()

    # Preface
    preface_heading = doc.add_paragraph('Preface')
    preface_heading.runs[0].bold = True
    preface_heading.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    preface_paragraphs = [
        "When I first sat down to write this book, I found myself staring at a blank screen for what felt like hours. The cursor blinked patiently, as if waiting for me to articulate something I had known for years but struggled to express in words. The art of coding is not simply about making computers do what you want — it is about creating something beautiful, maintainable, and comprehensible.",
        "Over the course of fifteen years in software engineering, I have worked on systems ranging from embedded microcontrollers with 8KB of flash memory to distributed platforms handling billions of requests per day. What I have discovered, time and again, is that the principles of good coding transcend scale. A well-named variable matters in a ten-line script just as much as in a million-line enterprise application.",
        "This book is intended for programmers at every level. Whether you are just starting out or have been writing code for decades, I believe there is something here for you. The chapters are designed to build upon one another, but each can also stand alone as a reference on a particular topic.",
        "I owe a debt of gratitude to the many colleagues, mentors, and students who shaped my thinking over the years. In particular, I want to thank Dr. Maria Santos at Stanford University, whose course on software design changed the way I think about systems. I also want to thank the entire team at Northbridge Publishing for their patience and support throughout this project.",
    ]

    for text in preface_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Chapter 1
    ch1_heading = doc.add_paragraph('Chapter 1: The Foundation of Clean Code')
    ch1_heading.runs[0].bold = True
    ch1_heading.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    ch1_intro = doc.add_paragraph('Section 1.1 — What Is Clean Code?')
    ch1_intro.runs[0].bold = True
    ch1_intro.runs[0].font.size = Pt(13)

    ch1_paragraphs = [
        "Ask ten programmers to define clean code and you will likely receive ten different answers. Some will tell you it is code that runs fast. Others will say it is code with no bugs. Still others will point to brevity or elegance. All of these answers contain a grain of truth, but none of them captures the full picture.",
        "Clean code, at its core, is code that communicates its intent clearly to the next reader — and that reader is often you, six months from now. It is code that does what it says, says what it does, and does not surprise the person maintaining it at 2am on a Tuesday when the production server is failing.",
        "Robert C. Martin, in his seminal work on software craftsmanship, described clean code as code that could be read and enhanced by a developer other than its original author. This definition emphasizes two key qualities: readability and modifiability. These two qualities are deeply intertwined. Code that is easy to read is usually easy to modify, because the reader can understand what each part does without needing to trace through complex logic trees.",
        "Consider the following example. Imagine you come across a function in a legacy codebase. It is two hundred lines long, takes five parameters with names like 'd', 'tmp', and 'x2', and contains nested loops three levels deep. Even if this function runs correctly and efficiently, it is the opposite of clean code. Every modification risks introducing a bug, because the logic is so opaque that even the original author may have forgotten what it does.",
    ]

    for text in ch1_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    ch1_section2 = doc.add_paragraph('Section 1.2 — The Cost of Messy Code')
    ch1_section2.runs[0].bold = True
    ch1_section2.runs[0].font.size = Pt(13)

    ch1_s2_paragraphs = [
        "The cost of messy code is not immediately obvious. When a team is racing to meet a deadline, cutting corners on code quality feels like a reasonable trade-off. The feature ships, the client is happy, and the sprint ends with a sense of accomplishment. But the debt accumulates.",
        "Technical debt — the accumulated cost of shortcuts and compromises made during development — behaves much like financial debt. A small amount is manageable and sometimes even strategic. But when it grows unchecked, it begins to compound. Teams find themselves spending more and more time working around old problems rather than building new features.",
        "Studies by Capers Jones, a leading researcher in software productivity, suggest that poor code quality is responsible for roughly 30 to 40 percent of all software project failures. The remainder are attributed to requirements issues, management problems, and resource constraints. This means that the code itself — the artifact that programmers produce — is the single largest controllable factor in whether a software project succeeds or fails.",
        "The good news is that clean code is not magic. It is a skill that can be taught, practiced, and improved over time. The chapters that follow will explore specific techniques and principles that experienced programmers use to keep their code clean, their systems maintainable, and their teams productive.",
    ]

    for text in ch1_s2_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Chapter 2
    ch2_heading = doc.add_paragraph('Chapter 2: Naming Things Well')
    ch2_heading.runs[0].bold = True
    ch2_heading.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    ch2_intro_section = doc.add_paragraph('Section 2.1 — The Power of a Good Name')
    ch2_intro_section.runs[0].bold = True
    ch2_intro_section.runs[0].font.size = Pt(13)

    ch2_paragraphs = [
        "Phil Karlton, a legendary programmer at Netscape, famously said that there are only two hard problems in computer science: cache invalidation and naming things. He was only half joking. Choosing the right name for a variable, function, or class is one of the most impactful decisions a programmer makes.",
        "A good name is like a good label on a filing cabinet. It tells you exactly what is inside without requiring you to open the drawer. A bad name is like a label that says 'stuff' — technically accurate, but completely useless for navigation.",
        "Consider a variable named 'data'. This name conveys essentially no information. What kind of data? Data about what? From where? Compare it to a variable named 'monthlyRevenueByRegion'. Suddenly, the purpose and structure of the variable are immediately clear, even to someone who has never seen the code before.",
        "Intention-revealing names are the first and most important principle of clean code. Every name you choose should answer three questions: what this thing is, why it exists, and how it is used. If a name requires a comment to explain it, the name has failed.",
    ]

    for text in ch2_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    ch2_section2 = doc.add_paragraph('Section 2.2 — Common Naming Mistakes')
    ch2_section2.runs[0].bold = True
    ch2_section2.runs[0].font.size = Pt(13)

    ch2_s2_paragraphs = [
        "The most common naming mistake is using abbreviations and acronyms. While it is tempting to save keystrokes by writing 'usr' instead of 'user' or 'mgr' instead of 'manager', this habit creates code that is harder to read and search. Modern IDEs handle long names with ease, and there is no practical benefit to abbreviation in most contexts.",
        "Another common mistake is using misleading names. A variable called 'accountList' that is actually an array, not a linked list, will confuse every programmer who reads it. Similarly, a function called 'getUser' that actually modifies the database is a trap waiting to spring on an unsuspecting developer.",
        "Noise words are another source of confusion. Names like 'theCustomer', 'aProduct', or 'customerInfo' are redundant. The articles and suffixes add no information. Simply 'customer' or 'product' is cleaner and more direct.",
        "Finally, using different names for the same concept — or the same name for different concepts — creates unnecessary cognitive load. If you use 'fetch', 'retrieve', and 'get' interchangeably across your codebase, readers must constantly wonder whether these operations have subtle differences. Pick one term and stick to it.",
    ]

    for text in ch2_s2_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Chapter 3
    ch3_heading = doc.add_paragraph('Chapter 3: Functions and Their Secrets')
    ch3_heading.runs[0].bold = True
    ch3_heading.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    ch3_intro_section = doc.add_paragraph('Section 3.1 — The Single Responsibility Principle')
    ch3_intro_section.runs[0].bold = True
    ch3_intro_section.runs[0].font.size = Pt(13)

    ch3_paragraphs = [
        "Functions are the primary unit of organization in most programming languages. They are how we group related operations, hide complexity, and create reusable components. But functions can also become the primary source of confusion if they are not written carefully.",
        "The single responsibility principle, often abbreviated as SRP, states that a function should do one thing, and do it well. This sounds simple in theory, but in practice it requires constant discipline. It is always tempting to add just one more operation to a function that is already partially doing what you need.",
        "A function that reads a file, parses its contents, validates the data, and writes results to a database is doing at least four things. Each of those four operations is a potential failure point, and each requires a different kind of expertise to understand and test. Breaking this function into four smaller functions — each responsible for one operation — makes the code dramatically easier to understand, test, and modify.",
        "The question of how small is small enough is often debated among programmers. Some advocate for functions of no more than ten lines. Others are comfortable with twenty or thirty lines if the logic is cohesive. The exact number matters less than the underlying principle: each function should have a clear, describable purpose that can be expressed in a single sentence without the word 'and'.",
    ]

    for text in ch3_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    ch3_section2 = doc.add_paragraph('Section 3.2 — Function Arguments')
    ch3_section2.runs[0].bold = True
    ch3_section2.runs[0].font.size = Pt(13)

    ch3_s2_paragraphs = [
        "The number of arguments a function takes is inversely correlated with its readability and testability. A function with no arguments — a niladic function — is the easiest to call, test, and understand. A function with one argument — a monadic function — is nearly as simple. The complexity grows with each additional argument.",
        "When a function requires three or more arguments, this is often a sign that the arguments belong together in a data structure. A function that takes 'firstName', 'lastName', 'email', and 'phoneNumber' as separate parameters would be cleaner if it accepted a single 'UserContact' object instead.",
        "Flag arguments — boolean parameters that control which of two paths the function executes — are a particular red flag. A function that behaves differently based on a boolean argument is really two functions in disguise. The boolean should be eliminated and replaced with two clearly named functions.",
        "Output arguments are also problematic. Functions that modify their arguments rather than returning values create surprising behavior. Programmers expect arguments to be input, not output. If a function must modify state, it should do so on an object it is called on, not on a parameter passed to it.",
    ]

    for text in ch3_s2_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Chapter 4
    ch4_heading = doc.add_paragraph('Chapter 4: Comments — When and Why')
    ch4_heading.runs[0].bold = True
    ch4_heading.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    ch4_intro_section = doc.add_paragraph('Section 4.1 — The Trouble with Comments')
    ch4_intro_section.runs[0].bold = True
    ch4_intro_section.runs[0].font.size = Pt(13)

    ch4_paragraphs = [
        "Comments occupy an interesting place in the pantheon of programming opinions. Some developers view them as essential documentation, a gift to future readers. Others view them as admissions of defeat — evidence that the code is not clear enough to speak for itself. Both perspectives contain wisdom.",
        "The trouble with comments is that they have a tendency to become lies over time. Code changes; comments frequently do not. A comment that accurately described the behavior of a function when it was written may be completely misleading three refactoring cycles later. Code that contradicts its comments is worse than code with no comments at all.",
        "There is also a seductive quality to commenting bad code rather than fixing it. When faced with a confusing function, the path of least resistance is to add a comment explaining what it does. But the real solution is to rewrite the function so that no comment is needed. Every comment that exists because the code is unclear is a failure — a failure to write code that communicates directly.",
        "This does not mean comments should be avoided entirely. There are legitimate uses for comments, and a codebase with no comments at all can be just as frustrating as one drowning in outdated documentation. The key is to use comments for the right reasons.",
    ]

    for text in ch4_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    ch4_section2 = doc.add_paragraph('Section 4.2 — When Comments Add Value')
    ch4_section2.runs[0].bold = True
    ch4_section2.runs[0].font.size = Pt(13)

    ch4_s2_paragraphs = [
        "Legal comments — copyright notices, license headers — are often required by organizational policy and are entirely appropriate. They are not explaining code; they are providing legal context.",
        "Explanatory comments are valuable when code must implement a non-obvious algorithm or work around a documented bug in an external library. A comment that says 'We use this particular hash function because SHA-256 triggers a known bug in the Oracle JDBC driver version 11.x' saves future developers hours of investigation.",
        "Warning comments that explain why a particular approach was taken — or why an obvious alternative was deliberately avoided — are also genuinely useful. 'Do not use a thread pool here; this operation must run on the main thread due to UI framework constraints' is the kind of comment that prevents an expensive mistake.",
        "TODO comments serve a legitimate purpose when used to mark known limitations or planned improvements. But they should be treated as technical debt, not as permanent features. A codebase that has not had a TODO resolved in five years has a cultural problem, not a technical one.",
    ]

    for text in ch4_s2_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Chapter 5
    ch5_heading = doc.add_paragraph('Chapter 5: Error Handling Done Right')
    ch5_heading.runs[0].bold = True
    ch5_heading.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    ch5_intro_section = doc.add_paragraph('Section 5.1 — Error Handling as First-Class Concern')
    ch5_intro_section.runs[0].bold = True
    ch5_intro_section.runs[0].font.size = Pt(13)

    ch5_paragraphs = [
        "Error handling is often treated as an afterthought in software development. The happy path — the sequence of operations that succeeds as expected — gets the most attention during design and testing. Error cases are handled hastily, if at all. This is a serious mistake.",
        "Real-world software runs in an environment full of unexpected conditions. Networks fail. Disks fill up. APIs return unexpected data. Users do things that developers never anticipated. A system that handles errors gracefully is not just more reliable — it is also more trustworthy and easier to debug.",
        "The first principle of good error handling is: handle errors where you can do something useful about them. A low-level database function that fails to execute a query cannot know whether the appropriate response is to retry, to return a default value, or to abort the entire operation. That decision belongs at a higher level of abstraction.",
        "Exceptions and error codes represent two different philosophies of error communication. Error codes require the caller to check return values and propagate errors explicitly. Exceptions allow errors to travel up the call stack automatically until they reach a handler. Each approach has its place, but mixing them inconsistently within a single codebase creates confusion and bugs.",
    ]

    for text in ch5_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Chapter 6
    ch6_heading = doc.add_paragraph('Chapter 6: Unit Testing Philosophy')
    ch6_heading.runs[0].bold = True
    ch6_heading.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    ch6_intro_section = doc.add_paragraph('Section 6.1 — Why We Test')
    ch6_intro_section.runs[0].bold = True
    ch6_intro_section.runs[0].font.size = Pt(13)

    ch6_paragraphs = [
        "Unit testing has become one of the most widely accepted practices in software engineering, yet it remains one of the most inconsistently applied. Many development teams write tests; far fewer write tests that actually serve their intended purpose.",
        "The primary purpose of unit testing is not to verify that code works today. It is to verify that code still works tomorrow — after a refactoring, after a dependency update, after a new team member makes a change they thought was safe. Tests are a safety net that enables confident change.",
        "This distinction matters enormously in practice. Tests that are tightly coupled to implementation details — that break whenever the code is refactored even if the behavior is unchanged — provide no safety net at all. They create noise, erode trust in the test suite, and eventually get deleted. Tests should verify behavior, not implementation.",
        "The test-driven development (TDD) movement takes this philosophy to its logical conclusion. By writing tests before writing implementation code, TDD forces developers to think about the desired behavior first and the implementation second. This often leads to cleaner APIs and more modular code, because testable code tends to be well-structured code.",
    ]

    for text in ch6_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Chapter 7
    ch7_heading = doc.add_paragraph('Chapter 7: The Art of Refactoring')
    ch7_heading.runs[0].bold = True
    ch7_heading.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    ch7_intro_section = doc.add_paragraph('Section 7.1 — Refactoring Without Fear')
    ch7_intro_section.runs[0].bold = True
    ch7_intro_section.runs[0].font.size = Pt(13)

    ch7_paragraphs = [
        "Refactoring — the practice of improving the internal structure of code without changing its external behavior — is one of the most important skills a programmer can develop. It is also one of the most feared, particularly in legacy codebases where the consequences of change are unpredictable.",
        "The fear of refactoring is almost always a symptom of insufficient test coverage. When a codebase has comprehensive automated tests, refactoring is not risky — it is safe. The tests tell you immediately if a change has broken something. Without tests, every refactoring is a gamble.",
        "Martin Fowler's catalog of refactoring techniques, first published in 1999, remains the definitive reference on the subject. It describes over a hundred specific refactorings — from simple operations like renaming a variable to complex structural changes like replacing a type code with a class hierarchy. Each refactoring is described as a series of safe, incremental steps that preserve the observable behavior of the code.",
        "The most important insight from Fowler's work is that refactoring is not a special activity that happens occasionally. It is a continuous practice, woven into the everyday work of programming. Every time you work in a piece of code, you should leave it a little cleaner than you found it. This principle — sometimes called the Boy Scout Rule — prevents the accumulation of technical debt over time.",
    ]

    for text in ch7_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Chapter 8
    ch8_heading = doc.add_paragraph('Chapter 8: Design Patterns in Practice')
    ch8_heading.runs[0].bold = True
    ch8_heading.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    ch8_intro_section = doc.add_paragraph('Section 8.1 — Patterns as a Shared Vocabulary')
    ch8_intro_section.runs[0].bold = True
    ch8_intro_section.runs[0].font.size = Pt(13)

    ch8_paragraphs = [
        "Design patterns are reusable solutions to commonly occurring problems in software design. They were popularized by the book 'Design Patterns: Elements of Reusable Object-Oriented Software', published in 1994 by four authors who became collectively known as the Gang of Four.",
        "The value of design patterns is not primarily in the solutions themselves — experienced programmers would likely arrive at similar solutions independently. The value is in the shared vocabulary they provide. When a team member says 'we should use an observer pattern here', every other team member immediately understands the proposed structure, its trade-offs, and how to implement it.",
        "The twenty-three patterns described by the Gang of Four are divided into three categories: creational patterns (which deal with object creation), structural patterns (which deal with object composition), and behavioral patterns (which deal with algorithms and responsibility assignment).",
        "Critics of design patterns sometimes argue that they are workarounds for limitations of object-oriented languages, and that more expressive languages make many patterns unnecessary. There is some truth to this. The strategy pattern, for example, is simply a first-class function in a language that supports functional programming. But even in expressive languages, the vocabulary of patterns remains useful for communication and documentation.",
    ]

    for text in ch8_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Chapter 9
    ch9_heading = doc.add_paragraph('Chapter 9: Performance and Optimization')
    ch9_heading.runs[0].bold = True
    ch9_heading.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    ch9_intro_section = doc.add_paragraph('Section 9.1 — Measure Before You Optimize')
    ch9_intro_section.runs[0].bold = True
    ch9_intro_section.runs[0].font.size = Pt(13)

    ch9_paragraphs = [
        "Donald Knuth's famous observation that premature optimization is the root of all evil in programming has been quoted so often that it has almost lost its meaning. But the underlying insight remains as relevant as ever: optimizing code before you know where the performance bottleneck is will, at best, waste your time and, at worst, make the code worse.",
        "Modern profiling tools make it straightforward to identify where a program spends most of its time. In virtually every non-trivial program, performance is dominated by a small fraction of the code — often 10 percent of the code consumes 90 percent of the execution time. Optimizing anything outside that hot path has no measurable impact on real-world performance.",
        "The first step in any performance optimization effort should be to establish a baseline. Without a measurement of current performance, there is no way to know whether an optimization has actually helped. Tools like perf on Linux, Instruments on macOS, and Visual Studio's performance profiler on Windows provide detailed execution traces that make hot paths visible.",
        "Algorithmic efficiency is almost always more important than micro-optimizations. Replacing an O(n²) algorithm with an O(n log n) alternative will outperform any amount of loop unrolling or branch prediction optimization for large inputs. Before micro-optimizing, make sure you have chosen the right algorithm.",
    ]

    for text in ch9_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Chapter 10
    ch10_heading = doc.add_paragraph('Chapter 10: The Human Side of Coding')
    ch10_heading.runs[0].bold = True
    ch10_heading.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    ch10_intro_section = doc.add_paragraph('Section 10.1 — Code Is Communication')
    ch10_intro_section.runs[0].bold = True
    ch10_intro_section.runs[0].font.size = Pt(13)

    ch10_paragraphs = [
        "Throughout this book, we have focused on the technical aspects of clean code: naming, structure, testing, and performance. But there is a dimension of software development that transcends all of these technical concerns: the human dimension.",
        "Software is not written by machines for machines. It is written by people for people — with machines merely as the medium. The primary audience for your code is not the computer, which will execute it blindly regardless of how it is written. The primary audience is the human being who will read, maintain, and extend it.",
        "This perspective changes how we think about every decision in software development. We do not write tests because machines require them. We write tests because tests communicate the intended behavior of our code to other developers, and because they enable the kind of fearless change that keeps codebases healthy over time.",
        "We choose meaningful names not because compilers care what variables are called — they absolutely do not — but because names are the primary way we communicate our mental model of a problem to other programmers. A well-named variable tells a story; a poorly named one requires the reader to reverse-engineer the story from the surrounding code.",
    ]

    for text in ch10_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    ch10_section2 = doc.add_paragraph('Section 10.2 — The Craft Mindset')
    ch10_section2.runs[0].bold = True
    ch10_section2.runs[0].font.size = Pt(13)

    ch10_s2_paragraphs = [
        "The most effective programmers I have known all share a common attitude toward their work: they think of programming as a craft, not merely a job. They take pride in their code the way a skilled carpenter takes pride in a well-made chair. They are not satisfied with code that merely works — they want code that is elegant, clear, and maintainable.",
        "This craft mindset has practical consequences. Craftspeople invest in their tools — their editors, debuggers, and build systems. They practice deliberately, reading other people's code, studying new languages and paradigms, and working on side projects that stretch their abilities. They mentor others, because teaching is one of the most effective ways to deepen understanding.",
        "The craft mindset also means taking responsibility for quality. A craftsperson who produces shoddy work does not blame the deadline or the requirements. They find a way to meet the deadline without compromising on essential quality — or they have an honest conversation about trade-offs with their client or manager.",
        "Software development has been called an engineering discipline, a science, a form of mathematics, and an art. It is, in truth, all of these things and none of them entirely. But the metaphor of craft captures something that the others miss: the combination of technical skill and personal investment that distinguishes truly excellent programmers from merely competent ones.",
    ]

    for text in ch10_s2_paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    # Closing paragraph
    conclusion_para = doc.add_paragraph()
    run = conclusion_para.add_run("The art of coding is a lifelong pursuit. The principles in this book are not destinations but signposts on a continuous journey. I hope they serve you well.")
    run.italic = True

    doc.add_paragraph()

    # Acknowledgments
    ack_heading = doc.add_paragraph('Acknowledgments')
    ack_heading.runs[0].bold = True
    ack_heading.runs[0].font.size = Pt(16)

    ack_paras = [
        "I am grateful to the entire team at Northbridge Publishing, especially my editor, Helena Kowalski, whose thoughtful feedback made this book substantially better than it would have been otherwise.",
        "My colleagues at Meridian Software, past and present, have been an invaluable source of ideas and inspiration. In particular, I want to thank David Okonkwo, Priya Nair, and Thomas Bergström for many enlightening discussions about software design over the years.",
        "Finally, and most importantly, I want to thank my family — my wife, Claire, and my children, Sofia and Eli — for their patience and support throughout the writing process. This book is dedicated to them.",
    ]

    for text in ack_paras:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
