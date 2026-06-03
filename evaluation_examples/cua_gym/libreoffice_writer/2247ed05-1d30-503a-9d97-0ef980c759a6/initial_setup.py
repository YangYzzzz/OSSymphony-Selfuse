"""
Initial Setup: Set up different headers for odd and even pages in a user manual
Task ID: writer_rd_021
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_021'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_body_content(doc, chapter_num, chapter_title, paragraphs_needed):
    """Add realistic body content for a chapter to fill pages."""
    # Chapter heading
    heading = doc.add_heading(chapter_title, level=1)

    # Realistic content per chapter
    chapter_content = {
        1: [
            "This user manual provides comprehensive instructions for operating the XR-500 Industrial Control System. The system is designed for manufacturing environments requiring precise temperature, pressure, and flow control across multiple production lines.",
            "The XR-500 was developed by Meridian Technologies in partnership with leading industrial automation experts. Since its introduction in 2022, it has been deployed in over 350 facilities worldwide, managing critical processes in chemical manufacturing, food processing, and pharmaceutical production.",
            "Before proceeding with installation, ensure that all prerequisite conditions have been met. The facility must have a stable 480V three-phase power supply with dedicated grounding. Network infrastructure should support Ethernet/IP with a minimum bandwidth of 100 Mbps for real-time data acquisition.",
            "Safety is paramount when working with the XR-500. All personnel must complete the mandatory safety training program (minimum 16 hours) before accessing the system. Personal protective equipment including safety glasses, ESD wrist straps, and steel-toed boots must be worn at all times in the control room.",
            "The control system architecture consists of three primary layers: the field device layer (sensors and actuators), the control layer (PLCs and distributed I/O), and the supervisory layer (HMI workstations and data historians). Each layer communicates through redundant network connections to ensure continuous operation.",
        ],
        2: [
            "The installation process begins with site preparation. Ensure the control cabinet location meets the environmental specifications: ambient temperature between 5 and 40 degrees Celsius, relative humidity below 85% (non-condensing), and adequate ventilation for heat dissipation.",
            "Mount the main control panel on a stable, vibration-free surface using the provided M12 anchor bolts. The panel weighs approximately 185 kg and requires a minimum of two trained technicians for safe installation. Verify that the mounting surface can support at least 250 kg of static load.",
            "Connect the power supply cables according to the wiring diagram in Appendix B. The main power input requires 3-phase 480V AC at 60 Hz with a dedicated 100A circuit breaker. The UPS backup connection uses single-phase 240V AC and provides up to 30 minutes of emergency operation.",
            "Network configuration follows a star topology with the main switch located in the control cabinet. Each field device connects through shielded Cat6a cables with a maximum run length of 90 meters. Fiber optic connections are recommended for distances exceeding 50 meters or in areas with high electromagnetic interference.",
            "After completing the physical installation, perform the initial power-on sequence. Turn on the UPS first, wait for stable output (green LED on front panel), then energize the main circuit breaker. The system will perform a self-diagnostic routine lasting approximately 3 minutes.",
        ],
        3: [
            "The operator interface consists of a 21.5-inch industrial touchscreen display running the Meridian HMI software. The main dashboard provides real-time overview of all process variables, alarm status, and system health indicators. Navigation between screens uses the tab bar at the bottom of the display.",
            "Process monitoring is organized into five primary screens: Overview, Temperature Control, Pressure Management, Flow Regulation, and Quality Metrics. Each screen displays current values, setpoints, trends, and alarm limits for the relevant parameters. Historical data can be accessed through the trend viewer.",
            "The alarm management system uses a four-tier priority scheme: Critical (red), High (orange), Medium (yellow), and Low (blue). Critical alarms require immediate operator acknowledgment and cannot be silenced until the underlying condition is resolved. The alarm history log retains all events for a minimum of 90 days.",
            "User access is controlled through role-based authentication. Four default roles are configured: Viewer (read-only access), Operator (process control), Engineer (configuration changes), and Administrator (full system access). Each user must have a unique login credential, and passwords must be changed every 90 days.",
            "The reporting module generates daily, weekly, and monthly production reports automatically. Reports include key performance indicators such as Overall Equipment Effectiveness (OEE), Mean Time Between Failures (MTBF), and process capability indices (Cpk). Custom reports can be created using the built-in report designer.",
        ],
        4: [
            "Preventive maintenance is essential for reliable system operation. The recommended maintenance schedule includes daily visual inspections, weekly calibration checks, monthly component testing, and annual comprehensive overhauls. All maintenance activities must be documented in the maintenance log.",
            "Daily inspections should verify that all status indicator LEDs are showing normal (green) state, cooling fans are operating, and no unusual sounds or odors are present. Check the UPS battery status and verify that the backup power test completed successfully during the overnight cycle.",
            "Sensor calibration should be performed monthly or whenever measurement drift exceeds 0.5% of the full scale range. The calibration procedure uses NIST-traceable reference standards and follows the two-point calibration method described in the Meridian Calibration Guide (Document MCG-2024-R3).",
            "Replace the air filters in the control cabinet every three months or more frequently in dusty environments. Use only Meridian-approved filters (Part Number MF-2250) to ensure proper airflow and dust protection. Operating with clogged filters can cause overheating and premature component failure.",
            "The annual overhaul includes comprehensive testing of all safety interlocks, verification of redundant system failover, replacement of UPS batteries (recommended every 3 years), and firmware updates for all programmable devices. Schedule the overhaul during a planned production shutdown.",
        ],
        5: [
            "This chapter covers common troubleshooting procedures for the XR-500 system. Before attempting any troubleshooting, ensure you have the appropriate access level (Engineer or Administrator) and have reviewed the relevant safety procedures in Chapter 1.",
            "Communication failures between the HMI and PLCs are the most frequently reported issues. Start by checking the network switch port LEDs for link activity. If the link LED is off, verify cable connections at both ends. If the link is active but communication fails, check the IP address configuration and subnet mask settings.",
            "Temperature control loop oscillation can be caused by several factors: improperly tuned PID parameters, sensor drift, or actuator mechanical issues. Use the built-in auto-tune function (accessible from the Engineering screen) to recalculate PID values. If oscillation persists after auto-tuning, inspect the control valve for mechanical binding.",
            "Intermittent alarm flooding often indicates a process variable operating near its alarm limit. Review the alarm setpoints and consider implementing alarm deadbands to prevent rapid alarm cycling. The recommended deadband is 2% of the measurement span for analog alarms.",
            "If the system fails to start after a power outage, first verify that the UPS has recharged sufficiently (minimum 50% capacity shown on the UPS display). Then perform the standard power-on sequence described in Chapter 2. If the self-diagnostic routine reports errors, consult the error code table in Appendix C.",
        ],
        6: [
            "The XR-500 supports integration with enterprise systems through OPC UA, MQTT, and REST API interfaces. OPC UA provides secure, reliable communication for real-time data exchange with SCADA systems and data historians. The built-in OPC UA server supports up to 500 simultaneous client connections.",
            "MQTT integration enables lightweight messaging for IoT applications and cloud connectivity. The system publishes process data to configurable MQTT topics at rates up to 1000 messages per second. Quality of Service levels 0, 1, and 2 are supported for different reliability requirements.",
            "The REST API provides HTTP-based access to system data and configuration. All API endpoints require OAuth 2.0 authentication tokens obtained from the system's identity server. Rate limiting is enforced at 100 requests per second per client to prevent system overload.",
            "Database integration supports direct writing to SQL Server, PostgreSQL, and Oracle databases. The data historian module buffers up to 72 hours of process data locally in case of database connectivity loss. When connectivity is restored, buffered data is automatically uploaded in chronological order.",
            "For custom integration scenarios, the XR-500 provides a Python SDK that simplifies development of third-party applications. The SDK includes classes for reading process values, writing setpoints, managing alarms, and subscribing to real-time data streams. Sample code and documentation are available on the Meridian Developer Portal.",
        ],
    }

    content = chapter_content.get(chapter_num, chapter_content[1])
    for i in range(paragraphs_needed):
        text = content[i % len(content)]
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(8)
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.name = 'Calibri'


def create_initial():
    doc = Document()

    # Set up page margins
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Set up a simple header: 'User Manual' on ALL pages (same left/right)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ''
    run = hp.add_run('User Manual')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Title page
    title_para = doc.add_heading('XR-500 Industrial Control System', level=0)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('User Manual')
    run.font.size = Pt(24)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    version_para = doc.add_paragraph()
    version_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = version_para.add_run('Version 4.2.1 | March 2025')
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author_para.add_run('Meridian Technologies, Inc.')
    run.font.size = Pt(12)
    run.font.name = 'Calibri'

    doc.add_page_break()

    # Table of Contents placeholder
    toc_heading = doc.add_heading('Table of Contents', level=1)
    chapters = [
        ('1', 'Introduction and Safety'),
        ('2', 'Installation and Setup'),
        ('3', 'Operation and Interface'),
        ('4', 'Maintenance and Calibration'),
        ('5', 'Troubleshooting'),
        ('6', 'System Integration'),
    ]
    for num, title in chapters:
        toc_entry = doc.add_paragraph()
        run = toc_entry.add_run(f'Chapter {num}: {title}')
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    doc.add_page_break()

    # 6 chapters - distribute content to fill ~30 pages
    # Each chapter needs roughly 4-5 pages worth of paragraphs
    chapter_titles = [t for _, t in chapters]
    for i, title in enumerate(chapter_titles, 1):
        full_title = f'Chapter {i}: {title}'
        # Add enough paragraphs per chapter (~20 paragraphs each for ~4-5 pages)
        add_body_content(doc, i, full_title, paragraphs_needed=20)
        # Add page break between chapters (except after last)
        if i < 6:
            doc.add_page_break()

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
