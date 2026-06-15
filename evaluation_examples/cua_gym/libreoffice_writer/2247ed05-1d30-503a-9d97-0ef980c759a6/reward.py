"""
Reward Script: Different headers for odd and even pages in Writer document
Task ID: writer_rd_021
Domain: libreoffice_writer
Scoring:
  Component 1: evenAndOddHeaders enabled in document settings (0.2 pts)
  Component 2: Even page headers = 'User Manual', left-aligned, italic 10pt (0.25 pts)
  Component 3: Odd page headers = chapter titles, right-aligned, italic 10pt (0.3 pts)
  Component 4: Both odd and even headers have thin bottom border (0.25 pts)
"""

import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_021'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: evenAndOddHeaders is enabled in document settings (0.2 points)
    # Initial state has this OFF; golden has it ON.
    try:
        settings_elem = doc.settings._element
        eaoh = settings_elem.find(qn('w:evenAndOddHeaders'))
        if eaoh is not None:
            print(f"PASS: Component 1 — evenAndOddHeaders is enabled (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 1 — evenAndOddHeaders not found in document settings")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Even page headers contain 'User Manual', left-aligned, italic 10pt (0.25 points)
    # Initial state has NO even page header content (linked_to_previous=True).
    # Golden state has even headers with 'User Manual' in each section.
    try:
        sections = doc.sections
        even_header_pass_count = 0
        even_header_total = 0

        for i, sec in enumerate(sections):
            even_header = sec.even_page_header
            # Skip if linked to previous (means no custom even header)
            if even_header.is_linked_to_previous:
                continue

            even_header_total += 1
            paras = even_header.paragraphs
            if not paras:
                continue

            para = paras[0]
            text = para.text.strip()
            alignment = para.paragraph_format.alignment

            # Check text contains 'User Manual'
            text_ok = 'User Manual' in text
            # Check alignment is LEFT (or None which defaults to LEFT)
            align_ok = alignment in (WD_PARAGRAPH_ALIGNMENT.LEFT, None, 0)
            # Check italic and size on runs
            italic_ok = False
            size_ok = False
            for run in para.runs:
                if run.text.strip():
                    if run.font.italic:
                        italic_ok = True
                    if run.font.size and run.font.size == Pt(10):
                        size_ok = True

            if text_ok and align_ok and italic_ok and size_ok:
                even_header_pass_count += 1

        if even_header_total > 0 and even_header_pass_count == even_header_total:
            print(f"PASS: Component 2 — All {even_header_total} even headers have 'User Manual' left-aligned italic 10pt (0.25 pts)")
            total_score += 0.25
        elif even_header_total > 0 and even_header_pass_count > 0:
            partial = 0.25 * (even_header_pass_count / even_header_total)
            print(f"PARTIAL: Component 2 — {even_header_pass_count}/{even_header_total} even headers correct ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No valid even headers found (total sections with custom even header: {even_header_total})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Odd page headers contain chapter titles, right-aligned, italic 10pt (0.3 points)
    # Initial state has a single header 'User Manual' LEFT aligned, NOT italic.
    # Golden state has chapter-specific titles RIGHT aligned, italic.
    try:
        sections = doc.sections
        odd_header_pass_count = 0
        odd_header_total = 0

        for i, sec in enumerate(sections):
            header = sec.header
            if header.is_linked_to_previous:
                continue

            odd_header_total += 1
            paras = header.paragraphs
            if not paras:
                continue

            para = paras[0]
            text = para.text.strip()
            alignment = para.paragraph_format.alignment

            # The odd header should be RIGHT aligned
            align_ok = (alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT or alignment == 2)
            # Should NOT be the default 'User Manual' only (must be a chapter title)
            # Chapter titles start with 'Chapter' or contain specific content
            # The key distinction: it should NOT just be 'User Manual' (initial state)
            text_ok = len(text) > 0 and text != 'User Manual'
            # Check italic and 10pt
            italic_ok = False
            size_ok = False
            for run in para.runs:
                if run.text.strip():
                    if run.font.italic:
                        italic_ok = True
                    if run.font.size and run.font.size == Pt(10):
                        size_ok = True

            if text_ok and align_ok and italic_ok and size_ok:
                odd_header_pass_count += 1

        if odd_header_total > 0 and odd_header_pass_count == odd_header_total:
            print(f"PASS: Component 3 — All {odd_header_total} odd headers have chapter titles right-aligned italic 10pt (0.3 pts)")
            total_score += 0.3
        elif odd_header_total > 0 and odd_header_pass_count > 0:
            partial = 0.3 * (odd_header_pass_count / odd_header_total)
            print(f"PARTIAL: Component 3 — {odd_header_pass_count}/{odd_header_total} odd headers correct ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No valid odd headers with chapter titles found (sections checked: {odd_header_total})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Both odd and even headers have a thin bottom border line (0.25 points)
    # Initial state has NO border on the header.
    # Golden state has bottom border on all headers.
    try:
        sections = doc.sections
        border_pass_count = 0
        border_total = 0

        for i, sec in enumerate(sections):
            # Check odd header border
            header = sec.header
            if not header.is_linked_to_previous and header.paragraphs:
                border_total += 1
                para = header.paragraphs[0]
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    pBdr = pPr.find(qn('w:pBdr'))
                    if pBdr is not None:
                        bottom = pBdr.find(qn('w:bottom'))
                        if bottom is not None:
                            val = bottom.get(qn('w:val'))
                            if val and val != 'none':
                                border_pass_count += 1

            # Check even header border
            even_header = sec.even_page_header
            if not even_header.is_linked_to_previous and even_header.paragraphs:
                border_total += 1
                para = even_header.paragraphs[0]
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    pBdr = pPr.find(qn('w:pBdr'))
                    if pBdr is not None:
                        bottom = pBdr.find(qn('w:bottom'))
                        if bottom is not None:
                            val = bottom.get(qn('w:val'))
                            if val and val != 'none':
                                border_pass_count += 1

        if border_total > 0 and border_pass_count == border_total:
            print(f"PASS: Component 4 — All {border_total} headers have bottom border (0.25 pts)")
            total_score += 0.25
        elif border_total > 0 and border_pass_count > 0:
            partial = 0.25 * (border_pass_count / border_total)
            print(f"PARTIAL: Component 4 — {border_pass_count}/{border_total} headers have bottom border ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — No headers with bottom borders found (checked: {border_total})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
