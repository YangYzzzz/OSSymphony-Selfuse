"""
Initial Setup: Legal memorandum document with citation (no footnotes)
Task ID: writer_legal_008
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title ---
    title = doc.add_heading('MEMORANDUM OF LAW IN SUPPORT OF MOTION FOR SUMMARY JUDGMENT', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Case caption ---
    caption = doc.add_paragraph()
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = caption.add_run('Anderson Manufacturing Corp. v. Pacific Northwest Industries, LLC')
    run.bold = True
    run.font.size = Pt(13)

    case_no = doc.add_paragraph()
    case_no.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = case_no.add_run('Case No. 2024-CV-03847')
    run.font.size = Pt(11)

    court = doc.add_paragraph()
    court.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = court.add_run('United States District Court, Western District of Washington')
    run.font.size = Pt(11)

    doc.add_paragraph()  # blank line

    # --- I. Introduction ---
    intro_heading = doc.add_heading('I. INTRODUCTION', level=2)
    intro_p1 = doc.add_paragraph()
    intro_p1.paragraph_format.first_line_indent = Inches(0.5)
    intro_p1.add_run(
        'Plaintiff Anderson Manufacturing Corp. ("Anderson") respectfully submits this '
        'Memorandum of Law in support of its Motion for Summary Judgment pursuant to '
        'Federal Rule of Civil Procedure 56. As demonstrated below, there is no genuine '
        'dispute as to any material fact, and Anderson is entitled to judgment as a matter of law.'
    )

    intro_p2 = doc.add_paragraph()
    intro_p2.paragraph_format.first_line_indent = Inches(0.5)
    intro_p2.add_run(
        'This action arises from Defendant Pacific Northwest Industries, LLC\'s ("PNI") '
        'breach of a supply agreement dated March 15, 2023. Anderson seeks damages in the '
        'amount of $2,450,000 representing lost profits and consequential damages resulting '
        'from PNI\'s failure to deliver conforming goods within the contractually specified timeframe.'
    )

    # --- II. Statement of Facts ---
    doc.add_heading('II. STATEMENT OF UNDISPUTED MATERIAL FACTS', level=2)
    facts_p1 = doc.add_paragraph()
    facts_p1.paragraph_format.first_line_indent = Inches(0.5)
    facts_p1.add_run(
        'On March 15, 2023, Anderson and PNI entered into a Master Supply Agreement '
        '("the Agreement") under which PNI agreed to deliver 15,000 units of industrial-grade '
        'titanium fasteners conforming to ASTM F468 specifications. The Agreement specified '
        'delivery in three equal installments on April 30, June 30, and August 31, 2023. '
        '(Declaration of Robert Chen, Exh. A at \u00b6\u00b6 3-5.)'
    )

    facts_p2 = doc.add_paragraph()
    facts_p2.paragraph_format.first_line_indent = Inches(0.5)
    facts_p2.add_run(
        'PNI delivered the first installment of 5,000 units on May 12, 2023, twelve days '
        'after the contractual deadline. Quality inspection revealed that approximately 1,200 '
        'units (24%) failed to meet the specified tensile strength requirements. '
        '(Chen Decl., Exh. B; Expert Report of Dr. Lisa Nakamura at 8-12.)'
    )

    facts_p3 = doc.add_paragraph()
    facts_p3.paragraph_format.first_line_indent = Inches(0.5)
    facts_p3.add_run(
        'PNI failed to deliver the second and third installments entirely. On September 15, '
        '2023, PNI\'s Vice President of Operations, Margaret Sullivan, sent an email to Anderson\'s '
        'procurement director acknowledging that PNI "cannot fulfill the remaining obligations '
        'under the Agreement due to unexpected capacity constraints." (Chen Decl., Exh. C.)'
    )

    facts_p4 = doc.add_paragraph()
    facts_p4.paragraph_format.first_line_indent = Inches(0.5)
    facts_p4.add_run(
        'As a direct result of PNI\'s breach, Anderson was forced to source replacement '
        'fasteners from Meridian Metals Corporation at a 38% cost premium, and suffered '
        'a 47-day production delay on its contract with Boeing Defense Systems. '
        '(Declaration of Karen Okafor, \u00b6\u00b6 7-14.)'
    )

    # --- III. Argument ---
    doc.add_heading('III. ARGUMENT', level=2)
    arg_p1 = doc.add_paragraph()
    arg_p1.paragraph_format.first_line_indent = Inches(0.5)
    arg_p1.add_run(
        'Summary judgment is appropriate when "there is no genuine dispute as to any material '
        'fact and the movant is entitled to judgment as a matter of law." Fed. R. Civ. P. 56(a). '
        'The moving party bears the initial burden of demonstrating the absence of a genuine '
        'issue of material fact. Celotex Corp. v. Catrett, 477 U.S. 317, 323 (1986). Once that '
        'burden is met, the non-moving party must set forth specific facts showing a genuine '
        'issue for trial. Matsushita Elec. Indus. Co. v. Zenith Radio Corp., 475 U.S. 574, 587 (1986).'
    )

    # Second paragraph of Argument - contains the target citation
    arg_p2 = doc.add_paragraph()
    arg_p2.paragraph_format.first_line_indent = Inches(0.5)
    arg_p2.add_run(
        'The undisputed facts establish that PNI materially breached the Agreement. Under '
        'the Uniform Commercial Code, a buyer may recover damages for non-delivery or '
        'repudiation measured by the difference between the market price and the contract '
        'price. U.C.C. \u00a7 2-713. Moreover, as the Supreme Court held in '
        'Smith v. Jones, 456 U.S. 789 (2019), '
        'the standard for evaluating breach of commercial supply agreements requires '
        'consideration of both the objective terms of the contract and the commercially '
        'reasonable expectations of the parties at the time of formation.'
    )

    arg_p3 = doc.add_paragraph()
    arg_p3.paragraph_format.first_line_indent = Inches(0.5)
    arg_p3.add_run(
        'Here, PNI\'s own correspondence unequivocally admits its inability to perform. '
        'The September 15 email from Ms. Sullivan constitutes a clear anticipatory repudiation '
        'of the remaining obligations under the Agreement. See Hochster v. De La Tour, '
        '2 El. & Bl. 678 (1853); Restatement (Second) of Contracts \u00a7 253 (1981). '
        'No reasonable trier of fact could conclude otherwise.'
    )

    arg_p4 = doc.add_paragraph()
    arg_p4.paragraph_format.first_line_indent = Inches(0.5)
    arg_p4.add_run(
        'Anderson\'s damages are also established as a matter of law. The cost differential '
        'between the contract price ($18.50/unit) and the replacement price ($25.53/unit) '
        'for the 10,000 undelivered units totals $70,300. Additional consequential damages '
        'of $2,379,700 are supported by the expert testimony of Dr. James Whitfield, who '
        'calculated Anderson\'s lost profits from the Boeing production delay with reasonable '
        'certainty. (Whitfield Expert Report at 15-28.)'
    )

    # --- IV. Conclusion ---
    doc.add_heading('IV. CONCLUSION', level=2)
    concl_p1 = doc.add_paragraph()
    concl_p1.paragraph_format.first_line_indent = Inches(0.5)
    concl_p1.add_run(
        'For the foregoing reasons, Anderson Manufacturing Corp. respectfully requests that '
        'this Court grant its Motion for Summary Judgment and enter judgment in its favor '
        'in the amount of $2,450,000, plus pre-judgment interest, costs, and attorneys\' fees '
        'as permitted under Section 12.4 of the Agreement.'
    )

    doc.add_paragraph()  # blank line

    # Signature block
    sig = doc.add_paragraph('Respectfully submitted,')
    sig.paragraph_format.space_before = Pt(24)

    doc.add_paragraph()
    sig_name = doc.add_paragraph()
    run = sig_name.add_run('_________________________________')
    sig_name2 = doc.add_paragraph()
    run = sig_name2.add_run('Katherine R. Blackwell, Esq.')
    run.bold = True
    bar = doc.add_paragraph('Washington State Bar No. 38472')
    firm = doc.add_paragraph('Blackwell & Associates, PLLC')
    addr = doc.add_paragraph('1200 Fifth Avenue, Suite 3400')
    city = doc.add_paragraph('Seattle, WA 98101')
    phone = doc.add_paragraph('Tel: (206) 555-0142')
    email = doc.add_paragraph('kblackwell@blackwelllaw.com')

    date_line = doc.add_paragraph()
    date_line.paragraph_format.space_before = Pt(12)
    date_line.add_run('Dated: January 15, 2025')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
