"""
Initial Setup: Add company logo and text to header
Task ID: writer_fs_078
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_078'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
LOGO_PATH = f'{WORKDIR}/Desktop/logo.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_logo():
    """Create a simple 2cm x 2cm logo image at /home/user/Desktop/logo.png."""
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)
    # 2cm at 96dpi ~ 76 pixels
    size_px = 76
    img = Image.new('RGB', (size_px, size_px), color=(0, 51, 102))
    # Draw a simple "A" letter pattern for Acme Corp
    from PIL import ImageDraw, ImageFont
    draw = ImageDraw.Draw(img)
    # Draw a white circle
    draw.ellipse([8, 8, size_px - 8, size_px - 8], outline='white', width=3)
    # Draw "A" in center
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 36)
    except (IOError, OSError):
        font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), "A", font=font)
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    x = (size_px - text_w) // 2
    y = (size_px - text_h) // 2 - 4
    draw.text((x, y), "A", fill='white', font=font)
    img.save(LOGO_PATH, dpi=(96, 96))
    print(f'Logo created: {LOGO_PATH}')


def create_initial():
    doc = Document()

    # Enable header but leave it empty
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    # Clear any default content - ensure header is empty
    for para in header.paragraphs:
        para.text = ""

    # Set reasonable margins
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Body Content: A realistic business memo ---
    heading = doc.add_heading('Quarterly Business Review - Q1 2025', level=1)

    doc.add_paragraph('')  # spacer

    # Date and metadata
    meta = doc.add_paragraph()
    meta.add_run('Date: ').bold = True
    meta.add_run('March 28, 2025')

    dept = doc.add_paragraph()
    dept.add_run('Department: ').bold = True
    dept.add_run('Operations & Strategy')

    prepared = doc.add_paragraph()
    prepared.add_run('Prepared by: ').bold = True
    prepared.add_run('Sarah Chen, VP of Operations')

    doc.add_paragraph('')

    # Executive Summary
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'This report summarizes the key operational metrics and strategic initiatives '
        'for Acme Corporation during the first quarter of 2025. Overall revenue grew by '
        '12.3% year-over-year, driven primarily by strong performance in the Enterprise '
        'Solutions division and the successful launch of our new cloud platform.'
    )

    # Key Metrics
    doc.add_heading('Key Performance Metrics', level=2)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Metric', 'Q1 2025', 'Q4 2024', 'Change']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Revenue', '$14.2M', '$12.8M', '+10.9%'],
        ['New Customers', '347', '289', '+20.1%'],
        ['Customer Retention', '94.7%', '93.2%', '+1.5%'],
        ['Employee Count', '412', '398', '+3.5%'],
        ['Net Promoter Score', '72', '68', '+4 pts'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')

    # Strategic Initiatives
    doc.add_heading('Strategic Initiatives', level=2)
    doc.add_paragraph(
        'Cloud Platform Migration',
        style='List Bullet'
    )
    doc.add_paragraph(
        'The migration of legacy on-premise clients to our new cloud infrastructure '
        'is proceeding ahead of schedule. As of March 15, 2025, 68% of enterprise '
        'clients have completed the transition, with the remaining scheduled for Q2.'
    )

    doc.add_paragraph(
        'International Expansion',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Our APAC office in Singapore opened on February 3, 2025, with an initial '
        'team of 15 employees. Early traction has been promising, with 23 qualified '
        'leads generated in the first six weeks of operation.'
    )

    doc.add_paragraph(
        'Product Development',
        style='List Bullet'
    )
    doc.add_paragraph(
        'The R&D team delivered three major feature releases during Q1, including '
        'advanced analytics dashboards, automated reporting workflows, and enhanced '
        'API integration capabilities for third-party platforms.'
    )

    # Next Steps
    doc.add_heading('Next Steps', level=2)
    doc.add_paragraph('Complete remaining cloud migrations by end of Q2 2025.', style='List Number')
    doc.add_paragraph('Expand Singapore office to 30 employees by September 2025.', style='List Number')
    doc.add_paragraph('Launch beta program for AI-powered analytics module.', style='List Number')
    doc.add_paragraph('Begin planning for Series C funding round.', style='List Number')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


# Create the logo first
create_logo()

# Create the initial document
create_initial()
