"""
Initial Setup: Compare thesis draft with advisor's edited version
Task ID: writer_lec_067
Domain: libreoffice_writer
Creates two documents: thesis_draft.docx (original) and thesis_advisor_edits.docx (edited)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_067'
DRAFT_PATH = f'{WORKDIR}/thesis_draft.docx'
EDITS_PATH = f'{WORKDIR}/thesis_advisor_edits.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_draft():
    """Create the original thesis draft with standard formatting."""
    doc = Document()

    # --- Standard formatting for the draft ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Title
    title = doc.add_heading('The Impact of Machine Learning on Climate Modeling: '
                            'A Comprehensive Analysis', level=0)
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Author info
    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.space_after = Pt(6)
    run = author_para.add_run('Megan L. Torres')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    dept_para = doc.add_paragraph()
    dept_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept_para.paragraph_format.space_after = Pt(18)
    run = dept_para.add_run('Department of Environmental Science, Westfield University')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True

    # --- Chapter 1: Introduction ---
    h1 = doc.add_heading('Chapter 1: Introduction', level=1)
    for run in h1.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)

    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.5
    p1.paragraph_format.space_after = Pt(8)
    run = p1.add_run(
        'Climate modeling has undergone significant transformations over the past two decades. '
        'Traditional physics-based models, while foundational, often struggle to capture the '
        'nonlinear interactions between atmospheric, oceanic, and terrestrial systems. The '
        'emergence of machine learning techniques has opened new avenues for improving both the '
        'accuracy and computational efficiency of these models.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run(
        'This thesis examines the application of deep learning architectures, specifically '
        'convolutional neural networks and transformer models, to regional climate prediction. '
        'Our research focuses on the Pacific Northwest region, where complex topography and '
        'marine influences create particularly challenging prediction scenarios.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1.5
    p3.paragraph_format.space_after = Pt(8)
    run = p3.add_run(
        'The primary contributions of this work include: (1) a novel hybrid architecture that '
        'combines physics-informed constraints with data-driven learning, (2) a comprehensive '
        'evaluation framework for comparing ML-enhanced models against traditional approaches, '
        'and (3) an analysis of model interpretability in the context of climate science.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # --- Chapter 2: Literature Review ---
    h2 = doc.add_heading('Chapter 2: Literature Review', level=1)
    for run in h2.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)

    p4 = doc.add_paragraph()
    p4.paragraph_format.line_spacing = 1.5
    p4.paragraph_format.space_after = Pt(8)
    run = p4.add_run(
        'Early attempts to integrate statistical learning into climate models date back to the '
        'work of Hasselmann (1997), who proposed stochastic parameterization schemes. Subsequent '
        'developments by Reichstein et al. (2019) demonstrated the potential of deep learning for '
        'geoscientific applications, particularly in carbon cycle modeling and extreme weather '
        'event prediction.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p5 = doc.add_paragraph()
    p5.paragraph_format.line_spacing = 1.5
    p5.paragraph_format.space_after = Pt(8)
    run = p5.add_run(
        'The intersection of physics-based and data-driven approaches has been extensively '
        'explored by Beucler et al. (2021), who introduced physically-constrained neural networks '
        'for climate model emulation. Their work showed that incorporating conservation laws as '
        'hard constraints improved generalization performance by 34% on out-of-distribution samples.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p6 = doc.add_paragraph()
    p6.paragraph_format.line_spacing = 1.5
    p6.paragraph_format.space_after = Pt(8)
    run = p6.add_run(
        'Regional downscaling remains one of the most promising applications of ML in climate '
        'science. Vandal et al. (2017) developed DeepSD, a deep learning framework for statistical '
        'downscaling that outperformed traditional methods such as BCSD and ARRM on precipitation '
        'prediction tasks across the continental United States.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # --- Chapter 3: Methodology ---
    h3 = doc.add_heading('Chapter 3: Methodology', level=1)
    for run in h3.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)

    p7 = doc.add_paragraph()
    p7.paragraph_format.line_spacing = 1.5
    p7.paragraph_format.space_after = Pt(8)
    run = p7.add_run(
        'Our experimental framework consists of three main components: data preprocessing, '
        'model architecture design, and evaluation methodology. We utilize ERA5 reanalysis data '
        'from the European Centre for Medium-Range Weather Forecasts covering the period 1979-2023, '
        'with spatial resolution of 0.25 degrees.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p8 = doc.add_paragraph()
    p8.paragraph_format.line_spacing = 1.5
    p8.paragraph_format.space_after = Pt(8)
    run = p8.add_run(
        'The proposed ClimateTransformer architecture employs a modified vision transformer '
        'backbone with cross-attention mechanisms to integrate multiple atmospheric variables. '
        'Input features include temperature, pressure, humidity, wind speed, and geopotential '
        'height at 13 pressure levels, along with sea surface temperature and land surface '
        'properties as auxiliary inputs.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # --- Chapter 4: Results ---
    h4 = doc.add_heading('Chapter 4: Results', level=1)
    for run in h4.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)

    p9 = doc.add_paragraph()
    p9.paragraph_format.line_spacing = 1.5
    p9.paragraph_format.space_after = Pt(8)
    run = p9.add_run(
        'Our ClimateTransformer model achieved a root mean square error of 1.47 degrees Celsius '
        'for 72-hour temperature forecasts, representing a 23% improvement over the baseline '
        'numerical weather prediction model. For precipitation prediction, the model demonstrated '
        'a critical success index of 0.62 for events exceeding 25mm per day.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p10 = doc.add_paragraph()
    p10.paragraph_format.line_spacing = 1.5
    p10.paragraph_format.space_after = Pt(8)
    run = p10.add_run(
        'Ablation studies revealed that the physics-informed constraints contributed approximately '
        '15% of the overall performance gain, while the cross-attention mechanism for multi-variable '
        'integration accounted for an additional 8%. The remaining improvement was attributed to '
        'the increased model capacity and the transformer architecture\'s ability to capture '
        'long-range spatial dependencies.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.save(DRAFT_PATH)
    print(f'Draft created: {DRAFT_PATH}')


def create_advisor_edits():
    """Create the advisor's edited version with both formatting and content changes."""
    doc = Document()

    # --- FORMATTING CHANGE: Advisor uses Calibri instead of Times New Roman ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title - FORMATTING CHANGE: different size, blue color
    title = doc.add_heading('The Impact of Machine Learning on Climate Modeling: '
                            'A Comprehensive Analysis', level=0)
    for run in title.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Author info
    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.space_after = Pt(6)
    run = author_para.add_run('Megan L. Torres')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    dept_para = doc.add_paragraph()
    dept_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept_para.paragraph_format.space_after = Pt(18)
    run = dept_para.add_run('Department of Environmental Science, Westfield University')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.italic = True

    # --- Chapter 1: Introduction ---
    h1 = doc.add_heading('Chapter 1: Introduction', level=1)
    for run in h1.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(18)  # FORMATTING CHANGE: 18pt instead of 16pt

    # CONTENT CHANGE: Advisor rewrote the first paragraph
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 2.0  # FORMATTING CHANGE: double spacing
    p1.paragraph_format.space_after = Pt(10)  # FORMATTING CHANGE: more spacing
    run = p1.add_run(
        'Over the last twenty years, climate modeling has been fundamentally reshaped by advances '
        'in computational science. While physics-based models remain indispensable, they frequently '
        'fail to adequately represent the complex nonlinear feedbacks among atmospheric, oceanic, '
        'and terrestrial subsystems. Machine learning methods now offer promising pathways to '
        'enhance both predictive accuracy and computational throughput of these models.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 2.0  # FORMATTING CHANGE
    p2.paragraph_format.space_after = Pt(10)
    run = p2.add_run(
        'This thesis examines the application of deep learning architectures, specifically '
        'convolutional neural networks and transformer models, to regional climate prediction. '
        'Our research focuses on the Pacific Northwest region, where complex topography and '
        'marine influences create particularly challenging prediction scenarios.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # CONTENT CHANGE: Advisor added a sentence at the end of the contributions paragraph
    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 2.0  # FORMATTING CHANGE
    p3.paragraph_format.space_after = Pt(10)
    run = p3.add_run(
        'The primary contributions of this work include: (1) a novel hybrid architecture that '
        'combines physics-informed constraints with data-driven learning, (2) a comprehensive '
        'evaluation framework for comparing ML-enhanced models against traditional approaches, '
        'and (3) an analysis of model interpretability in the context of climate science. '
        'Additionally, we provide open-source implementations of all proposed methods to '
        'facilitate reproducibility.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # --- Chapter 2: Literature Review ---
    h2 = doc.add_heading('Chapter 2: Literature Review', level=1)
    for run in h2.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(18)  # FORMATTING CHANGE

    p4 = doc.add_paragraph()
    p4.paragraph_format.line_spacing = 2.0  # FORMATTING CHANGE
    p4.paragraph_format.space_after = Pt(10)
    run = p4.add_run(
        'Early attempts to integrate statistical learning into climate models date back to the '
        'work of Hasselmann (1997), who proposed stochastic parameterization schemes. Subsequent '
        'developments by Reichstein et al. (2019) demonstrated the potential of deep learning for '
        'geoscientific applications, particularly in carbon cycle modeling and extreme weather '
        'event prediction.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # CONTENT CHANGE: Advisor deleted the Beucler paragraph and replaced with different text
    p5 = doc.add_paragraph()
    p5.paragraph_format.line_spacing = 2.0
    p5.paragraph_format.space_after = Pt(10)
    run = p5.add_run(
        'More recently, foundation models such as Pangu-Weather (Bi et al., 2023) and GraphCast '
        '(Lam et al., 2023) have achieved remarkable results in medium-range weather forecasting, '
        'surpassing operational NWP systems on several standard benchmarks. These developments '
        'suggest a paradigm shift in how atmospheric prediction may be approached.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    p6 = doc.add_paragraph()
    p6.paragraph_format.line_spacing = 2.0  # FORMATTING CHANGE
    p6.paragraph_format.space_after = Pt(10)
    run = p6.add_run(
        'Regional downscaling remains one of the most promising applications of ML in climate '
        'science. Vandal et al. (2017) developed DeepSD, a deep learning framework for statistical '
        'downscaling that outperformed traditional methods such as BCSD and ARRM on precipitation '
        'prediction tasks across the continental United States.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # --- Chapter 3: Methodology ---
    h3 = doc.add_heading('Chapter 3: Methodology', level=1)
    for run in h3.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(18)  # FORMATTING CHANGE

    p7 = doc.add_paragraph()
    p7.paragraph_format.line_spacing = 2.0  # FORMATTING CHANGE
    p7.paragraph_format.space_after = Pt(10)
    run = p7.add_run(
        'Our experimental framework consists of three main components: data preprocessing, '
        'model architecture design, and evaluation methodology. We utilize ERA5 reanalysis data '
        'from the European Centre for Medium-Range Weather Forecasts covering the period 1979-2023, '
        'with spatial resolution of 0.25 degrees.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # CONTENT CHANGE: Advisor rewrote the architecture paragraph
    p8 = doc.add_paragraph()
    p8.paragraph_format.line_spacing = 2.0
    p8.paragraph_format.space_after = Pt(10)
    run = p8.add_run(
        'We introduce ClimateTransformer, a novel architecture based on the vision transformer '
        'paradigm that leverages multi-head cross-attention to jointly process heterogeneous '
        'atmospheric variables. The model ingests temperature, pressure, specific humidity, wind '
        'components, and geopotential height fields across 13 standard pressure levels, supplemented '
        'by ocean surface and land surface boundary conditions.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # --- Chapter 4: Results ---
    h4 = doc.add_heading('Chapter 4: Results', level=1)
    for run in h4.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(18)  # FORMATTING CHANGE

    p9 = doc.add_paragraph()
    p9.paragraph_format.line_spacing = 2.0  # FORMATTING CHANGE
    p9.paragraph_format.space_after = Pt(10)
    run = p9.add_run(
        'Our ClimateTransformer model achieved a root mean square error of 1.47 degrees Celsius '
        'for 72-hour temperature forecasts, representing a 23% improvement over the baseline '
        'numerical weather prediction model. For precipitation prediction, the model demonstrated '
        'a critical success index of 0.62 for events exceeding 25mm per day.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    p10 = doc.add_paragraph()
    p10.paragraph_format.line_spacing = 2.0  # FORMATTING CHANGE
    p10.paragraph_format.space_after = Pt(10)
    run = p10.add_run(
        'Ablation studies revealed that the physics-informed constraints contributed approximately '
        '15% of the overall performance gain, while the cross-attention mechanism for multi-variable '
        'integration accounted for an additional 8%. The remaining improvement was attributed to '
        'the increased model capacity and the transformer architecture\'s ability to capture '
        'long-range spatial dependencies.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    doc.save(EDITS_PATH)
    print(f'Advisor edits created: {EDITS_PATH}')


def main():
    create_draft()
    create_advisor_edits()

    # Launch LibreOffice Writer with the draft open
    launch_gui(f'libreoffice --writer "{DRAFT_PATH}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
