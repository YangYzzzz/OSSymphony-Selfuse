"""
Initial Setup: Create a Writer document with quoted passages in Default Paragraph Style.
Task ID: writer_bs_077
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_077'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    doc.add_heading("The Evolution of Modern Architecture", level=1)

    doc.add_paragraph(
        "Architecture has always reflected the cultural and technological "
        "aspirations of its time. From the towering cathedrals of medieval "
        "Europe to the sleek glass towers of contemporary cities, each era "
        "leaves its mark on the built environment."
    )

    doc.add_paragraph(
        "The modernist movement, which emerged in the early twentieth century, "
        "fundamentally changed how architects approached design. Rejecting "
        "ornamental excess, modernists embraced function, simplicity, and "
        "new materials like reinforced concrete and steel."
    )

    # A quoted passage (in default style, no Blockquote style)
    doc.add_paragraph(
        '"Architecture is the learned game, correct and magnificent, of forms '
        'assembled in the light." - Le Corbusier, Towards a New Architecture, 1923'
    )

    doc.add_paragraph(
        "Le Corbusier's vision influenced generations of architects who sought "
        "to strip buildings down to their essential elements. His five points "
        "of architecture became a manifesto for the International Style."
    )

    doc.add_heading("The Bauhaus Legacy", level=2)

    doc.add_paragraph(
        "Founded in 1919 by Walter Gropius in Weimar, Germany, the Bauhaus "
        "school merged fine arts with craftsmanship. Its curriculum emphasized "
        "experimentation with materials and a holistic approach to design that "
        "encompassed furniture, textiles, typography, and architecture."
    )

    # Another quoted passage
    doc.add_paragraph(
        '"The ultimate aim of all creative activity is the building. The decoration '
        'of buildings was once the noblest function of the fine arts, and the fine arts '
        'were indispensable to great architecture." - Walter Gropius, Bauhaus Manifesto, 1919'
    )

    doc.add_paragraph(
        "When the Nazis closed the Bauhaus in 1933, its teachers dispersed across "
        "the globe, carrying modernist principles to new shores. Ludwig Mies van der "
        "Rohe settled in Chicago, where he designed the iconic Farnsworth House and "
        "the Illinois Institute of Technology campus."
    )

    doc.add_heading("Postmodern Reactions", level=2)

    doc.add_paragraph(
        "By the 1960s, critics began to question modernism's rigidity. Robert Venturi's "
        "landmark text challenged the purist ethos with his famous declaration:"
    )

    # Another quoted passage
    doc.add_paragraph(
        '"Less is a bore." - Robert Venturi, Complexity and Contradiction in '
        'Architecture, 1966'
    )

    doc.add_paragraph(
        "Postmodernism reintroduced color, ornamentation, and historical references "
        "into architecture. Buildings by Michael Graves, Philip Johnson, and Charles "
        "Moore playfully mixed classical motifs with contemporary forms, sparking "
        "debate about taste, meaning, and the role of architecture in public life."
    )

    doc.add_paragraph(
        "Today, the conversation has moved beyond stylistic labels. Sustainability, "
        "digital fabrication, and parametric design are reshaping the discipline. "
        "Yet the fundamental questions posed by the modernists and their critics "
        "remain as relevant as ever: What should a building express? How should "
        "form relate to function? And who ultimately benefits from the spaces we create?"
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
