"""
Initial Setup: Meeting minutes document with 8 paragraphs
Task ID: wrpara_017
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_017'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Paragraph 1: Meeting header/introduction
    p1 = doc.add_paragraph()
    p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p1.add_run('Quarterly Operations Meeting Minutes - March 28, 2025')
    run.bold = True
    run.font.size = Pt(14)

    # Paragraph 2: Budget Review (part 1)
    p2 = doc.add_paragraph(
        'Budget Review: CFO Sarah Chen presented the Q1 financial summary. '
        'Total revenue reached $4.2 million, exceeding projections by 8%. '
        'Operating expenses were held at $3.1 million, resulting in a net margin of 26.2%. '
        'The marketing department requested an additional $150,000 for the digital campaign expansion.'
    )

    # Paragraph 3: Budget Review (part 2) - separate paragraph to be merged with p2
    p3 = doc.add_paragraph(
        'The finance team confirmed that discretionary spending remains under the quarterly cap of $500,000. '
        'Sarah recommended reallocating $75,000 from the travel budget to cover the marketing request, '
        'noting that remote client meetings have reduced travel costs by 40% since January.'
    )

    # Paragraph 4: Project Timeline Update
    p4 = doc.add_paragraph(
        'Project Timeline Update: Director of Engineering Marcus Johnson reported that the '
        'Atlas Platform migration is 72% complete, ahead of the April 30 deadline. '
        'The remaining modules include payment processing integration and user authentication upgrades. '
        'QA testing for Phase 2 is scheduled to begin on April 7, with a dedicated team of six engineers.'
    )

    # Paragraph 5: Staffing Updates (part 1)
    p5 = doc.add_paragraph(
        'Staffing Updates: HR Manager Priya Patel announced that 12 new hires started in March across '
        'three departments: Engineering (5), Customer Support (4), and Sales (3). '
        'The onboarding satisfaction score averaged 4.6 out of 5.0, the highest in company history.'
    )

    # Paragraph 6: Staffing Updates (part 2) - separate paragraph to be merged with p5
    p6 = doc.add_paragraph(
        'Two senior developer positions remain open with final-round interviews scheduled for next week. '
        'Priya also highlighted the upcoming leadership training program starting April 14, '
        'which will include 20 participants from mid-level management across all regional offices.'
    )

    # Paragraph 7: Action Items
    p7 = doc.add_paragraph(
        'Action Items: (1) Sarah Chen to submit the revised budget allocation proposal by April 4. '
        '(2) Marcus Johnson to provide a detailed QA test plan for Phase 2 by April 5. '
        '(3) Priya Patel to finalize the leadership training curriculum and distribute the schedule by April 10. '
        '(4) All department heads to submit Q2 objectives by April 11.'
    )

    # Paragraph 8: Closing
    p8 = doc.add_paragraph(
        'The meeting was adjourned at 3:45 PM. The next quarterly operations meeting is scheduled for '
        'June 27, 2025, at 2:00 PM in Conference Room B. Minutes recorded by Administrative Assistant David Kim.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
