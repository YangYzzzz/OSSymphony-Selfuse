"""
Initial Setup: multi_column.docx - document with body text, no text frames
Task ID: writer_obj_070
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_070'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/multi_column.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set up page margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("Quarterly Business Report")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Calibri"

    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run("Regional Performance Analysis — Q1 2025")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.name = "Calibri"
    subtitle_run.italic = True

    doc.add_paragraph()  # spacer

    # Section heading
    heading_para = doc.add_paragraph()
    heading_run = heading_para.add_run("Executive Summary")
    heading_run.bold = True
    heading_run.font.size = Pt(13)
    heading_run.font.name = "Calibri"

    # Body text paragraph 1
    body1 = doc.add_paragraph()
    body1_run = body1.add_run(
        "This report provides a comprehensive overview of our business performance "
        "across three key operational divisions during the first quarter of 2025. "
        "Overall revenue grew by 12.4% year-over-year, driven primarily by strong "
        "results in the Asia-Pacific and European markets."
    )
    body1_run.font.size = Pt(11)
    body1_run.font.name = "Calibri"

    # Body text paragraph 2
    body2 = doc.add_paragraph()
    body2_run = body2.add_run(
        "Customer acquisition costs declined by 8.7% as our digital marketing "
        "initiatives began to yield measurable efficiency improvements. The sales team "
        "successfully closed 47 enterprise contracts, representing a total contract value "
        "of $3.2 million — a record high for Q1 performance."
    )
    body2_run.font.size = Pt(11)
    body2_run.font.name = "Calibri"

    # Section heading 2
    heading2_para = doc.add_paragraph()
    heading2_run = heading2_para.add_run("Key Performance Indicators")
    heading2_run.bold = True
    heading2_run.font.size = Pt(13)
    heading2_run.font.name = "Calibri"

    # Body text paragraph 3
    body3 = doc.add_paragraph()
    body3_run = body3.add_run(
        "The following metrics represent our most critical success indicators for "
        "the reporting period. Product delivery timelines improved across all three "
        "divisions, with the logistics team achieving a 99.1% on-time delivery rate — "
        "exceeding our internal benchmark of 98.5% for the fourth consecutive quarter."
    )
    body3_run.font.size = Pt(11)
    body3_run.font.name = "Calibri"

    # Body text paragraph 4
    body4 = doc.add_paragraph()
    body4_run = body4.add_run(
        "Staff retention rates remained high at 94.3%, supported by the new employee "
        "wellness program launched in January 2025. Training completion rates across "
        "departments averaged 87.6%, with the engineering and product teams achieving "
        "100% completion of all mandatory certification modules."
    )
    body4_run.font.size = Pt(11)
    body4_run.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
