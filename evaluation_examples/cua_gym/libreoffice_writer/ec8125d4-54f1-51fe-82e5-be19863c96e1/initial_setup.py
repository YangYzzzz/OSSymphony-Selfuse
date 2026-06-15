"""
Initial Setup: Monthly Newsletter - Two Column Layout
Task ID: writer_obj_055
Domain: libreoffice_writer

Creates monthly_newsletter.docx with a title and body content but NO text boxes.
The agent task is to add the two linked text boxes.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_055'
OUTPUT = f'{WORKDIR}/monthly_newsletter.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set page margins for a standard newsletter
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # --- Title ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('Monthly Newsletter - March 2026')
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Spacing after title
    title_para.paragraph_format.space_after = Pt(12)

    # --- Subtitle/Date ---
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = sub_para.add_run('Volume 8, Issue 3  |  March 2026  |  Community Edition')
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    sub_para.paragraph_format.space_after = Pt(8)

    # --- Introduction paragraph ---
    intro_para = doc.add_paragraph()
    intro_run = intro_para.add_run(
        "Welcome to the March edition of our monthly newsletter. This month we bring you "
        "the latest updates from our teams across departments, highlights from recent events, "
        "and a look ahead at what's coming in April. We hope you find this issue informative "
        "and engaging. As always, your feedback helps us improve."
    )
    intro_run.font.size = Pt(11)
    intro_para.paragraph_format.space_after = Pt(10)

    # --- Section: From the Editor ---
    heading1 = doc.add_paragraph()
    h1_run = heading1.add_run("From the Editor's Desk")
    h1_run.bold = True
    h1_run.font.size = Pt(14)
    h1_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(6)

    para1 = doc.add_paragraph()
    p1_run = para1.add_run(
        "March has been an exciting month for our organization. We welcomed five new team members "
        "to the engineering and marketing departments, and our Q1 results have exceeded expectations. "
        "The strategic partnership announced last week opens up new opportunities for collaboration "
        "across international markets."
    )
    p1_run.font.size = Pt(11)
    para1.paragraph_format.space_after = Pt(8)

    para2 = doc.add_paragraph()
    p2_run = para2.add_run(
        "Our community outreach program reached over 2,000 households this quarter, providing "
        "digital literacy training and technical support. Volunteer coordinators Maria Santos and "
        "James Okonkwo led the initiative with remarkable dedication."
    )
    p2_run.font.size = Pt(11)
    para2.paragraph_format.space_after = Pt(8)

    # --- Section: Team Spotlight ---
    heading2 = doc.add_paragraph()
    h2_run = heading2.add_run("Team Spotlight: Engineering Excellence")
    h2_run.bold = True
    h2_run.font.size = Pt(14)
    h2_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)

    para3 = doc.add_paragraph()
    p3_run = para3.add_run(
        "This month's spotlight falls on the infrastructure team who successfully migrated our "
        "primary data center to the new cloud architecture ahead of schedule. Led by Senior Engineer "
        "Rachel Kim, the migration resulted in a 40% reduction in latency and a 25% decrease in "
        "operational costs. The project involved coordinating across three time zones and multiple vendors."
    )
    p3_run.font.size = Pt(11)
    para3.paragraph_format.space_after = Pt(8)

    para4 = doc.add_paragraph()
    p4_run = para4.add_run(
        "The team has also completed the rollout of the new monitoring dashboard, which provides "
        "real-time visibility into system performance. Early adopters report significant improvements "
        "in incident response times, down from 45 minutes to under 8 minutes on average."
    )
    p4_run.font.size = Pt(11)
    para4.paragraph_format.space_after = Pt(8)

    # --- Section: Upcoming Events ---
    heading3 = doc.add_paragraph()
    h3_run = heading3.add_run("Upcoming Events — April 2026")
    h3_run.bold = True
    h3_run.font.size = Pt(14)
    h3_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    heading3.paragraph_format.space_before = Pt(12)
    heading3.paragraph_format.space_after = Pt(6)

    events = [
        ("April 3", "All-Hands Meeting — Conference Room A, 10:00 AM"),
        ("April 7-9", "Regional Sales Conference — Grand Hotel, Downtown"),
        ("April 15", "Q1 Financial Review — Board Room, 2:00 PM"),
        ("April 22", "Earth Day Community Cleanup — City Park, 9:00 AM"),
        ("April 28", "Product Launch Webinar — Online, 3:00 PM"),
    ]

    for date, event in events:
        event_para = doc.add_paragraph()
        date_run = event_para.add_run(f"{date}: ")
        date_run.bold = True
        date_run.font.size = Pt(11)
        event_run = event_para.add_run(event)
        event_run.font.size = Pt(11)
        event_para.paragraph_format.space_after = Pt(4)

    # --- Closing note ---
    closing_para = doc.add_paragraph()
    c_run = closing_para.add_run(
        "\nThank you for reading. To submit content for next month's newsletter, please contact "
        "the editorial team by April 20th. We look forward to sharing more exciting news with you."
    )
    c_run.font.size = Pt(11)
    c_run.italic = True
    closing_para.paragraph_format.space_before = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
