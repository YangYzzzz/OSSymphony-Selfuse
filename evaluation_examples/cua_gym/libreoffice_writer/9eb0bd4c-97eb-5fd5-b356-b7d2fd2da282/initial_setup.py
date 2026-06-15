"""
Initial Setup: Research paper with bold key terms (to be changed to italic by agent)
Task ID: writer_edit_013
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_013'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_para_with_bold_term(doc, before_text, bold_term, after_text):
    """Helper: add a paragraph containing a bold term embedded in normal text."""
    para = doc.add_paragraph()
    if before_text:
        run_before = para.add_run(before_text)
        run_before.bold = False
        run_before.italic = False
    run_bold = para.add_run(bold_term)
    run_bold.bold = True
    run_bold.italic = False
    if after_text:
        run_after = para.add_run(after_text)
        run_after.bold = False
        run_after.italic = False
    return para


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Advances in Artificial Intelligence: A Comprehensive Review", level=0)

    # --- Abstract ---
    doc.add_heading("Abstract", level=1)
    para = doc.add_paragraph()
    run = para.add_run(
        "This paper presents a systematic review of recent developments in artificial intelligence, "
        "with a particular focus on learning paradigms that have transformed the field over the past decade. "
        "We examine key methods and their applications across various domains, including computer vision, "
        "natural language understanding, and autonomous systems."
    )
    run.bold = False

    # --- Page 1: Introduction ---
    doc.add_heading("1. Introduction", level=1)

    para1 = doc.add_paragraph()
    r1 = para1.add_run(
        "The field of artificial intelligence has witnessed unprecedented growth in recent years. "
        "Central to this progress is the paradigm of "
    )
    r1.bold = False
    r1_bold = para1.add_run("machine learning")
    r1_bold.bold = True
    r1_bold.italic = False
    r1_end = para1.add_run(
        ", which enables systems to learn from data rather than relying solely on explicitly programmed rules. "
        "This shift has enabled computers to perform tasks that were once thought to require human intelligence."
    )
    r1_end.bold = False

    para2 = doc.add_paragraph()
    r2 = para2.add_run(
        "Among the most significant breakthroughs has been the rise of "
    )
    r2.bold = False
    r2_bold = para2.add_run("neural networks")
    r2_bold.bold = True
    r2_bold.italic = False
    r2_end = para2.add_run(
        ", computational models loosely inspired by the structure of the human brain. "
        "These architectures consist of interconnected layers of artificial neurons capable of "
        "learning complex patterns from high-dimensional data."
    )
    r2_end.bold = False

    para3 = doc.add_paragraph(
        "The organization of this paper is as follows: Section 2 reviews foundational learning paradigms, "
        "Section 3 examines state-of-the-art architectures, Section 4 discusses applications, "
        "and Section 5 concludes with future research directions."
    )

    # Page break after introduction
    doc.add_page_break()

    # --- Page 2: Foundational Learning Paradigms ---
    doc.add_heading("2. Foundational Learning Paradigms", level=1)

    doc.add_heading("2.1 Deep Learning", level=2)

    para4 = doc.add_paragraph()
    r4 = para4.add_run(
        "The emergence of "
    )
    r4.bold = False
    r4_bold = para4.add_run("deep learning")
    r4_bold.bold = True
    r4_bold.italic = False
    r4_end = para4.add_run(
        " has been one of the most transformative developments in modern AI. "
        "By stacking multiple layers of non-linear transformations, deep architectures can "
        "learn hierarchical representations from raw data, surpassing human performance on "
        "benchmark tasks in image recognition and speech synthesis."
    )
    r4_end.bold = False

    para5 = doc.add_paragraph(
        "The success of these architectures hinges on access to large-scale datasets, "
        "powerful hardware accelerators, and refined optimization algorithms such as "
        "stochastic gradient descent with momentum."
    )

    doc.add_heading("2.2 Reinforcement Learning", level=2)

    para6 = doc.add_paragraph()
    r6 = para6.add_run(
        "Another important paradigm is "
    )
    r6.bold = False
    r6_bold = para6.add_run("reinforcement learning")
    r6_bold.bold = True
    r6_bold.italic = False
    r6_end = para6.add_run(
        ", where an agent learns to make decisions by interacting with an environment "
        "and receiving scalar reward signals. This approach has achieved remarkable results "
        "in complex domains such as board games, video games, and robotic control, "
        "demonstrating the potential for autonomous decision-making."
    )
    r6_end.bold = False

    para7 = doc.add_paragraph(
        "Policy gradient methods and value-based approaches such as Q-learning have been "
        "combined in actor-critic frameworks that address the challenges of high-dimensional "
        "continuous action spaces and sparse reward signals."
    )

    # Page break
    doc.add_page_break()

    # --- Page 3: Supervised and Unsupervised Learning ---
    doc.add_heading("3. Supervised and Unsupervised Approaches", level=1)

    doc.add_heading("3.1 Supervised Learning", level=2)

    para8 = doc.add_paragraph()
    r8 = para8.add_run(
        "Traditional "
    )
    r8.bold = False
    r8_bold = para8.add_run("supervised learning")
    r8_bold.bold = True
    r8_bold.italic = False
    r8_end = para8.add_run(
        " relies on labeled training examples to guide the learning process. "
        "A model is trained to map input features to output labels by minimizing a loss function "
        "over the training set. Generalization to unseen data is measured on held-out test sets. "
        "Classification, regression, and sequence labeling are canonical supervised tasks."
    )
    r8_end.bold = False

    para9 = doc.add_paragraph(
        "Support vector machines, decision trees, and logistic regression were dominant approaches "
        "before the deep learning era. These methods remain competitive on low-data regimes and "
        "structured tabular data where interpretability is required."
    )

    doc.add_heading("3.2 Unsupervised Learning", level=2)

    para10 = doc.add_paragraph()
    r10 = para10.add_run(
        "In contrast, "
    )
    r10.bold = False
    r10_bold = para10.add_run("unsupervised learning")
    r10_bold.bold = True
    r10_bold.italic = False
    r10_end = para10.add_run(
        " attempts to discover structure in data without access to labels. "
        "Clustering algorithms, dimensionality reduction techniques such as principal component analysis, "
        "and generative models like variational autoencoders and generative adversarial networks "
        "are prominent examples of this approach."
    )
    r10_end.bold = False

    para11 = doc.add_paragraph(
        "Self-supervised learning has recently emerged as a powerful variant, "
        "leveraging auxiliary tasks derived from the data itself to learn useful representations "
        "without requiring expensive manual annotation."
    )

    # Page break
    doc.add_page_break()

    # --- Page 4: Transfer Learning and NLP ---
    doc.add_heading("4. Advanced Techniques", level=1)

    doc.add_heading("4.1 Transfer Learning", level=2)

    para12 = doc.add_paragraph()
    r12 = para12.add_run(
        "The concept of "
    )
    r12.bold = False
    r12_bold = para12.add_run("transfer learning")
    r12_bold.bold = True
    r12_bold.italic = False
    r12_end = para12.add_run(
        " involves leveraging knowledge acquired from one task or domain "
        "to improve performance on a related but different task. "
        "Pre-trained models that have been trained on large corpora can be fine-tuned "
        "with relatively small amounts of task-specific data, significantly reducing "
        "training time and computational cost."
    )
    r12_end.bold = False

    para13 = doc.add_paragraph(
        "This approach has been particularly influential in computer vision, "
        "where ImageNet-pretrained convolutional neural networks serve as feature extractors "
        "for downstream tasks including medical image analysis and satellite imagery classification."
    )

    doc.add_heading("4.2 Natural Language Processing", level=2)

    para14 = doc.add_paragraph()
    r14 = para14.add_run(
        "The field of "
    )
    r14.bold = False
    r14_bold = para14.add_run("natural language processing")
    r14_bold.bold = True
    r14_bold.italic = False
    r14_end = para14.add_run(
        " has been revolutionized by transformer-based architectures. "
        "Large language models trained on web-scale corpora demonstrate strong zero-shot "
        "and few-shot capabilities across diverse linguistic tasks including translation, "
        "summarization, question answering, and code generation."
    )
    r14_end.bold = False

    para15 = doc.add_paragraph(
        "The attention mechanism, which allows models to weigh the relevance of different "
        "input tokens when producing each output, has become a fundamental building block "
        "in modern sequence-to-sequence architectures."
    )

    para16 = doc.add_paragraph(
        "Ethical considerations surrounding bias, fairness, and interpretability remain active "
        "research areas as these systems are deployed in high-stakes applications such as "
        "healthcare diagnostics, legal decision support, and financial risk assessment."
    )

    # Page break
    doc.add_page_break()

    # --- Page 5: Conclusion ---
    doc.add_heading("5. Conclusion", level=1)

    para17 = doc.add_paragraph(
        "This review has surveyed the principal paradigms and methodologies that underpin "
        "contemporary artificial intelligence research. The convergence of large-scale data, "
        "computational resources, and algorithmic innovation has produced systems capable of "
        "remarkable feats across perception, reasoning, and decision-making."
    )

    para18 = doc.add_paragraph(
        "Future research directions include improving sample efficiency, "
        "developing more interpretable models, addressing distributional shift, "
        "and designing systems that integrate multiple modalities in a coherent fashion. "
        "Cross-disciplinary collaboration with cognitive science, neuroscience, and ethics "
        "will be essential to ensure these technologies benefit society broadly."
    )

    # References
    doc.add_heading("References", level=1)

    refs = [
        "LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep learning. Nature, 521(7553), 436-444.",
        "Sutton, R. S., & Barto, A. G. (2018). Reinforcement learning: An introduction. MIT Press.",
        "Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.",
        "Vaswani, A., et al. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30.",
        "Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. NAACL.",
        "Silver, D., et al. (2016). Mastering the game of Go with deep neural networks and tree search. Nature, 529, 484-489.",
        "Radford, A., et al. (2019). Language models are unsupervised multitask learners. OpenAI Blog.",
        "Pan, S. J., & Yang, Q. (2010). A survey on transfer learning. IEEE Transactions on Knowledge and Data Engineering.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='List Number')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
