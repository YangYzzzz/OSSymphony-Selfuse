"""
Initial Setup: Budget line items document with spaces (no tab stops)
Task ID: writer_biz_049
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_049'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title heading
    heading = doc.add_heading('FY2025 Annual Operating Budget', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sub.add_run('Greenfield Dynamics Inc.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Blank separator
    doc.add_paragraph()

    # Section header
    section_hdr = doc.add_paragraph()
    run = section_hdr.add_run('Department: Marketing & Communications')
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # 10 budget line items - using spaces (NOT tabs) between name and amount
    budget_items = [
        ('Digital Advertising Campaigns', '$128,500.00'),
        ('Trade Show & Conference Fees', '$67,250.00'),
        ('Content Production Services', '$43,800.00'),
        ('Social Media Management Tools', '$18,900.00'),
        ('Brand Collateral Printing', '$22,350.00'),
        ('Market Research & Analytics', '$55,600.00'),
        ('Public Relations Retainer', '$84,000.00'),
        ('Website Maintenance & Hosting', '$31,450.00'),
        ('Employee Training & Development', '$15,750.00'),
        ('Contingency & Miscellaneous', '$12,400.00'),
    ]

    for item_name, amount in budget_items:
        para = doc.add_paragraph()
        # Use spaces to separate name and amount (no tabs, no tab stops)
        run = para.add_run(f'{item_name}          {amount}')
        run.font.size = Pt(11)
        run.font.name = 'Liberation Serif'

    # Summary line
    doc.add_paragraph()
    total_para = doc.add_paragraph()
    run = total_para.add_run('Total Budget Allocation          $480,000.00')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Liberation Serif'

    # Footer note
    doc.add_paragraph()
    note = doc.add_paragraph()
    run = note.add_run('Prepared by Finance Department — Q1 Review Cycle')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
