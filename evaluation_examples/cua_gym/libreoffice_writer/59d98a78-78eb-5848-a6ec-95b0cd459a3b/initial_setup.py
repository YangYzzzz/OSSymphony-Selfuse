"""
Initial Setup: Create a contacts document with an unsorted table
Task ID: writer_tm_006
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()
    doc.add_heading("Contacts", level=1)

    # 4 columns x 8 rows (1 header + 7 data rows), unsorted by Last Name
    table = doc.add_table(rows=8, cols=4)
    table.style = "Table Grid"

    # Headers
    headers = ['First Name', 'Last Name', 'Email', 'Phone']
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    # Data rows - deliberately NOT sorted by Last Name
    data = [
        ['Maria',   'Zhang',    'mzhang@mail.com',    '555-0101'],
        ['John',    'Adams',    'jadams@mail.com',     '555-0202'],
        ['Sara',    'Lee',      'slee@mail.com',       '555-0303'],
        ['Derek',   'Patel',    'dpatel@mail.com',     '555-0404'],
        ['Lisa',    'Novak',    'lnovak@mail.com',     '555-0505'],
        ['Tom',     'Garcia',   'tgarcia@mail.com',    '555-0606'],
        ['Emily',   'Brooks',   'ebrooks@mail.com',    '555-0707'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            table.cell(row_idx, col_idx).text = val

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
