"""
Reward Script: Sort contacts table by Last Name ascending
Task ID: writer_tm_006
Domain: libreoffice_writer
Scoring:
  Precondition: File loads, table has 8 rows x 4 cols, header row intact
  Component 1 (0.5): Last Name column sorted in ascending order
  Component 2 (0.5): Row data coherence — each row's data matches the expected contact
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_006'

# Expected data rows after sorting by Last Name ascending
# Each entry: (First Name, Last Name, Email, Phone)
EXPECTED_SORTED_ROWS = [
    ('John', 'Adams', 'jadams@mail.com', '555-0202'),
    ('Emily', 'Brooks', 'ebrooks@mail.com', '555-0707'),
    ('Tom', 'Garcia', 'tgarcia@mail.com', '555-0606'),
    ('Sara', 'Lee', 'slee@mail.com', '555-0303'),
    ('Lisa', 'Novak', 'lnovak@mail.com', '555-0505'),
    ('Derek', 'Patel', 'dpatel@mail.com', '555-0404'),
    ('Maria', 'Zhang', 'mzhang@mail.com', '555-0101'),
]

EXPECTED_HEADER = ['First Name', 'Last Name', 'Email', 'Phone']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document

    total_score = 0.0

    # Load document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: table exists with correct structure
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    if num_rows != 8 or num_cols != 4:
        print(f"FAIL: Expected 8x4 table, found {num_rows}x{num_cols}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: header row intact
    header = [table.cell(0, c).text.strip() for c in range(4)]
    if header != EXPECTED_HEADER:
        print(f"FAIL: Header row mismatch. Expected {EXPECTED_HEADER}, found {header}")
        print("REWARD: 0.0")
        return 0.0

    print(f"PRECONDITION: Table 8x4, header row correct")

    # Read actual data rows
    actual_rows = []
    for r in range(1, 8):
        row_data = tuple(table.cell(r, c).text.strip() for c in range(4))
        actual_rows.append(row_data)

    # Component 1: Last Name column sorted ascending (0.5 points)
    try:
        last_names = [row[1] for row in actual_rows]
        sorted_last_names = sorted(last_names, key=str.lower)
        if last_names == sorted_last_names:
            print(f"PASS: Component 1 — Last names sorted ascending: {last_names} (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Last names not sorted. Found: {last_names}, expected: {sorted_last_names}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Row data coherence (0.5 points)
    # Each data row must match the expected sorted contact data.
    # Only awarded if Component 1 passed (sort is correct), to avoid
    # accidental partial credit when rows happen to be in the right spot.
    if total_score >= 0.5:
        try:
            correct_rows = 0
            for i, (actual, expected) in enumerate(zip(actual_rows, EXPECTED_SORTED_ROWS)):
                if actual == expected:
                    correct_rows += 1
                else:
                    print(f"  Row {i+1} mismatch: expected {expected}, found {actual}")

            # Progressive: award partial credit per correct row
            row_score = (correct_rows / 7) * 0.5
            if correct_rows == 7:
                print(f"PASS: Component 2 — All 7 data rows have correct coherent data (0.5 pts)")
                total_score += 0.5
            elif correct_rows > 0:
                print(f"PARTIAL: Component 2 — {correct_rows}/7 rows correct ({row_score:.3f} pts)")
                total_score += row_score
            else:
                print(f"FAIL: Component 2 — No data rows match expected sorted order")
        except Exception as e:
            print(f"ERROR: Component 2 — {e}")
    else:
        print(f"SKIP: Component 2 — Skipped because sort order (Component 1) failed")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice state before verification
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state()
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
