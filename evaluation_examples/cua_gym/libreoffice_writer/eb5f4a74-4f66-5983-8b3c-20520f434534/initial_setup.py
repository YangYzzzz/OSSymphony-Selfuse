"""
Initial Setup: Delete all three comments present in the document.
Task ID: writer_struct_012
Domain: libreoffice_writer

Creates a 3-page legal contract document with three reviewer comments:
  - 'Clause too broad' on 'indemnification'
  - 'Add 30-day notice' on 'termination'
  - 'Cap at $1M' on 'liability'
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_012'
OUTPUT = f'{WORKDIR}/Desktop/contract_review.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_comment_to_run(doc, paragraph, run_index, comment_text, author="Reviewer", date="2025-03-01T09:00:00Z"):
    """
    Add a comment to a specific run in a paragraph using OOXML manipulation.
    run_index: 0-based index of the run in the paragraph.
    """
    # Get or create the comments part
    body = doc.element.body

    # Find or create the comments.xml part
    # We'll use the document's comments relationship
    comments_part = None
    for rel in doc.part.rels.values():
        if 'comments' in rel.reltype.lower():
            comments_part = rel.target_part
            break

    if comments_part is None:
        # Need to create comments part - use existing approach via XML
        pass

    # We'll manipulate XML directly
    # Get the current maximum comment ID
    comments_xml_str = None
    try:
        comments_xml_str = doc.part.comments_part.xml if hasattr(doc.part, 'comments_part') else None
    except Exception:
        pass

    return None


def build_contract_with_comments():
    """Build a 3-page legal contract with embedded comments via direct XML."""
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # -------- Page 1: Contract Header & Parties --------
    heading = doc.add_heading('SERVICE AGREEMENT', level=0)
    heading.alignment = 1  # CENTER

    doc.add_paragraph('')

    intro = doc.add_paragraph(
        'This Service Agreement ("Agreement") is entered into as of March 1, 2025 '
        '("Effective Date") by and between Meridian Solutions Inc., a Delaware corporation '
        '("Service Provider"), and Hartwell Enterprises LLC, a New York limited liability '
        'company ("Client").'
    )

    doc.add_heading('1. SERVICES', level=1)
    doc.add_paragraph(
        'Service Provider agrees to provide software development and consulting services '
        'as described in Exhibit A attached hereto ("Services"). Service Provider shall '
        'perform the Services in a professional and workmanlike manner consistent with '
        'industry standards. Any changes to the scope of Services must be agreed upon in '
        'writing by both parties prior to commencement.'
    )

    doc.add_heading('2. COMPENSATION', level=1)
    doc.add_paragraph(
        'Client agrees to pay Service Provider a monthly retainer of $12,500 (twelve thousand '
        'five hundred dollars) due on the first business day of each month. Late payments shall '
        'accrue interest at the rate of 1.5% per month. All fees are non-refundable once work '
        'has commenced on a given deliverable. Invoices shall be submitted via electronic mail '
        'to accounts@hartwell.com no later than the 25th of the preceding month.'
    )

    doc.add_heading('3. CONFIDENTIALITY', level=1)
    doc.add_paragraph(
        'Each party agrees to maintain the confidentiality of the other party\'s proprietary '
        'information and trade secrets. Confidential information shall not be disclosed to third '
        'parties without prior written consent. This obligation shall survive the termination or '
        'expiration of this Agreement for a period of five (5) years. "Confidential Information" '
        'means any information that is marked as confidential or that reasonably should be '
        'understood to be confidential given the nature of the information and the circumstances '
        'of disclosure.'
    )

    doc.add_page_break()

    # -------- Page 2: Indemnification, Termination --------
    doc.add_heading('4. INDEMNIFICATION', level=1)
    p_indem = doc.add_paragraph(
        'Client agrees to defend, indemnify, and hold harmless Service Provider and its officers, '
        'directors, employees, and agents from and against any and all claims, damages, losses, '
        'costs, and expenses (including reasonable attorneys\' fees) arising out of or relating to: '
        '(a) Client\'s use of the Services; (b) Client\'s breach of this Agreement; or (c) any '
        'third-party claim arising from content or data provided by Client. The term '
    )
    run_indem = p_indem.add_run('indemnification')
    p_indem.add_run(
        ' obligations herein shall apply regardless of the form of action, whether in contract, '
        'tort, strict liability, or otherwise, and shall survive the termination of this Agreement.'
    )

    doc.add_heading('5. TERM AND TERMINATION', level=1)
    p_term = doc.add_paragraph(
        'This Agreement shall commence on the Effective Date and continue for an initial term of '
        'twelve (12) months unless earlier terminated. Either party may terminate this Agreement '
        'for cause upon written notice if the other party materially breaches any provision and '
        'fails to cure such breach within thirty (30) days of written notice. For convenience, '
        'either party may initiate '
    )
    run_term = p_term.add_run('termination')
    p_term.add_run(
        ' by providing sixty (60) days\' advance written notice. Upon termination, Service '
        'Provider shall deliver all completed work product and Client shall pay all fees accrued '
        'through the effective date of termination.'
    )

    doc.add_heading('6. INTELLECTUAL PROPERTY', level=1)
    doc.add_paragraph(
        'All work product, inventions, developments, and intellectual property created by Service '
        'Provider in performing the Services ("Work Product") shall be owned exclusively by Client '
        'upon full payment of all fees. Service Provider hereby assigns all right, title, and '
        'interest in the Work Product to Client. Service Provider retains the right to use '
        'general skills, knowledge, and expertise developed in performing the Services in future '
        'engagements with other clients, provided such use does not disclose Client\'s Confidential '
        'Information.'
    )

    doc.add_page_break()

    # -------- Page 3: Liability, Dispute Resolution, Signatures --------
    doc.add_heading('7. LIMITATION OF LIABILITY', level=1)
    p_liab = doc.add_paragraph(
        'IN NO EVENT SHALL SERVICE PROVIDER BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, '
        'CONSEQUENTIAL, OR PUNITIVE DAMAGES ARISING OUT OF OR RELATED TO THIS AGREEMENT, '
        'REGARDLESS OF WHETHER SUCH DAMAGES WERE FORESEEABLE. The total aggregate '
    )
    run_liab = p_liab.add_run('liability')
    p_liab.add_run(
        ' of Service Provider under this Agreement shall not exceed the total fees paid by '
        'Client to Service Provider during the twelve (12) months immediately preceding the '
        'event giving rise to the claim. Some jurisdictions do not allow the exclusion or '
        'limitation of incidental or consequential damages, so the above limitation may not '
        'apply to Client.'
    )

    doc.add_heading('8. DISPUTE RESOLUTION', level=1)
    doc.add_paragraph(
        'The parties agree to attempt to resolve any dispute arising under this Agreement through '
        'good faith negotiation. If negotiation fails, the parties shall submit the dispute to '
        'non-binding mediation before a mutually agreed mediator. If mediation is unsuccessful, '
        'disputes shall be resolved by binding arbitration in New York, New York, under the '
        'American Arbitration Association Commercial Arbitration Rules. The prevailing party shall '
        'be entitled to recover reasonable attorneys\' fees and costs.'
    )

    doc.add_heading('9. GENERAL PROVISIONS', level=1)
    doc.add_paragraph(
        'This Agreement constitutes the entire agreement between the parties with respect to its '
        'subject matter and supersedes all prior agreements, representations, and understandings. '
        'This Agreement may not be amended except by a written instrument signed by both parties. '
        'If any provision of this Agreement is held to be invalid or unenforceable, such provision '
        'shall be modified to the minimum extent necessary to make it valid and enforceable. '
        'This Agreement shall be governed by the laws of the State of Delaware without regard to '
        'conflict of law principles.'
    )

    doc.add_paragraph('')
    doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first above written.')
    doc.add_paragraph('')

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.cell(0, 0).text = 'MERIDIAN SOLUTIONS INC.'
    sig_table.cell(0, 1).text = 'HARTWELL ENTERPRISES LLC'
    sig_table.cell(1, 0).text = 'By: _______________________\nName: Alexandra Rivera\nTitle: Chief Executive Officer\nDate: ___________________'
    sig_table.cell(1, 1).text = 'By: _______________________\nName: Theodore Blackwell\nTitle: Managing Member\nDate: ___________________'

    # Save first, then inject comments via lxml
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)

    # Now inject comments via direct XML manipulation
    inject_comments(OUTPUT, run_indem, run_term, run_liab)

    print(f'Initial file created: {OUTPUT}')


def inject_comments(filepath, run_indem, run_term, run_liab):
    """
    Re-open the saved docx and inject three comments via direct XML manipulation.
    We use python-docx to load, then manipulate the lxml tree.
    """
    from docx import Document as DocxDocument
    import zipfile
    import shutil
    import re

    # Load doc again
    doc2 = DocxDocument(filepath)

    # Define the W namespace
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def make_qn(tag):
        return f'{{{W_NS}}}{tag}'

    # We'll build comments.xml content
    comments_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
  xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex"
  xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:aink="http://schemas.microsoft.com/office/drawing/2016/ink"
  xmlns:am3d="http://schemas.microsoft.com/office/drawing/2017/model3d"
  xmlns:o="urn:schemas-microsoft-com:office:office"
  xmlns:oel="http://schemas.microsoft.com/office/2019/extlst"
  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
  xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
  xmlns:v="urn:schemas-microsoft-com:vml"
  xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
  xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
  xmlns:w10="urn:schemas-microsoft-com:office:word"
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
  xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
  xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex"
  xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid"
  xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml"
  xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash"
  xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex"
  xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
  xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
  xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
  xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
  mc:Ignorable="w14 w15 w16se w16cid w16 w16cex w16sdtdh wp14">
  <w:comment w:id="1" w:author="Reviewer" w:date="2025-03-01T09:00:00Z" w:initials="R">
    <w:p><w:pPr><w:pStyle w:val="CommentText"/></w:pPr><w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:annotationRef/></w:r><w:r><w:t>Clause too broad</w:t></w:r></w:p>
  </w:comment>
  <w:comment w:id="2" w:author="Reviewer" w:date="2025-03-01T10:00:00Z" w:initials="R">
    <w:p><w:pPr><w:pStyle w:val="CommentText"/></w:pPr><w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:annotationRef/></w:r><w:r><w:t>Add 30-day notice</w:t></w:r></w:p>
  </w:comment>
  <w:comment w:id="3" w:author="Reviewer" w:date="2025-03-01T11:00:00Z" w:initials="R">
    <w:p><w:pPr><w:pStyle w:val="CommentText"/></w:pPr><w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:annotationRef/></w:r><w:r><w:t>Cap at $1M</w:t></w:r></w:p>
  </w:comment>
</w:comments>'''

    # Now we need to:
    # 1. Find the runs with 'indemnification', 'termination', 'liability' in the body XML
    # 2. Wrap each with commentRangeStart / commentRangeEnd / bookmarkStart / bookmarkEnd
    # 3. Add the comments.xml to the zip
    # 4. Add the relationship to document.xml.rels

    # Load doc body XML
    body_elem = doc2.element.body

    # Helper to find a run containing specific text and wrap it with comment marks
    def wrap_run_with_comment(body_elem, search_text, comment_id):
        """Find a run containing search_text and wrap it with comment range markers."""
        ns = W_NS
        for para in body_elem.iter(make_qn('p')):
            for run in para.iter(make_qn('r')):
                t_elem = run.find(make_qn('t'))
                if t_elem is not None and t_elem.text and search_text in t_elem.text:
                    # Found the run - insert comment range start before run
                    parent = run.getparent()
                    idx = list(parent).index(run)

                    # commentRangeStart
                    crs = OxmlElement('w:commentRangeStart')
                    crs.set(make_qn('id'), str(comment_id))
                    parent.insert(idx, crs)

                    # commentRangeEnd after run
                    cre = OxmlElement('w:commentRangeEnd')
                    cre.set(make_qn('id'), str(comment_id))
                    parent.insert(idx + 2, cre)

                    # commentReference run after commentRangeEnd
                    ref_run = OxmlElement('w:r')
                    ref_rpr = OxmlElement('w:rPr')
                    ref_rstyle = OxmlElement('w:rStyle')
                    ref_rstyle.set(make_qn('val'), 'CommentReference')
                    ref_rpr.append(ref_rstyle)
                    ref_run.append(ref_rpr)
                    ann_ref = OxmlElement('w:commentReference')
                    ann_ref.set(make_qn('id'), str(comment_id))
                    ref_run.append(ann_ref)
                    parent.insert(idx + 3, ref_run)
                    return True
        return False

    # Wrap the three target words
    wrap_run_with_comment(body_elem, 'indemnification', 1)
    wrap_run_with_comment(body_elem, 'termination', 2)
    wrap_run_with_comment(body_elem, 'liability', 3)

    # Save the modified document
    doc2.save(filepath)

    # Now inject comments.xml into the zip
    import zipfile
    import io

    # Read the docx as zip
    with zipfile.ZipFile(filepath, 'r') as zin:
        names = zin.namelist()
        files = {}
        for name in names:
            files[name] = zin.read(name)

    # Modify document.xml.rels to add comments relationship
    rels_key = 'word/_rels/document.xml.rels'
    if rels_key in files:
        rels_content = files[rels_key].decode('utf-8')
        if 'comments' not in rels_content.lower():
            # Add the comments relationship
            comments_rel = '<Relationship Id="rIdComments" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>'
            rels_content = rels_content.replace('</Relationships>', comments_rel + '</Relationships>')
            files[rels_key] = rels_content.encode('utf-8')

    # Modify [Content_Types].xml to add comments content type
    ct_key = '[Content_Types].xml'
    if ct_key in files:
        ct_content = files[ct_key].decode('utf-8')
        if 'comments' not in ct_content.lower():
            comments_ct = '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
            ct_content = ct_content.replace('</Types>', comments_ct + '</Types>')
            files[ct_key] = ct_content.encode('utf-8')

    # Add comments.xml
    files['word/comments.xml'] = comments_xml.encode('utf-8')

    # Write back to docx
    with zipfile.ZipFile(filepath, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)

    print(f'Comments injected into: {filepath}')


def create_initial():
    build_contract_with_comments()

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
