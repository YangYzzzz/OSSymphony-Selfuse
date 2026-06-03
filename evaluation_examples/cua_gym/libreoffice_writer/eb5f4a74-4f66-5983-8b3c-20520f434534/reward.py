"""
Reward Script: Delete all three comments from contract_review.docx
Task ID: writer_struct_012
Domain: libreoffice_writer
Scoring:
  Component 1: All comments are deleted from the document (no comment elements in CommentsPart) — 0.5 pts
  Component 2: Comment range markers removed from document body (no commentRangeStart/End/Reference) — 0.3 pts
  Component 3: Document text is preserved (paragraph count and key section headings intact) — 0.2 pts
"""

import os
import lxml.etree as etree

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_012'

FILE_PATH = f'{WORKDIR}/Desktop/contract_review.docx'

# Expected section headings that must remain after comments are deleted
EXPECTED_HEADINGS = [
    'SERVICE AGREEMENT',
    '1. SERVICES',
    '2. COMPENSATION',
    '3. CONFIDENTIALITY',
    '4. INDEMNIFICATION',
    '5. TERM AND TERMINATION',
    '6. INTELLECTUAL PROPERTY',
    '7. LIMITATION OF LIABILITY',
    '8. DISPUTE RESOLUTION',
    '9. GENERAL PROVISIONS',
]

COMMENTS_REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'


def verify_task(file_path):
    """
    Verify task completion: all three comments have been deleted from contract_review.docx.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — precondition gate
    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All comments are deleted from CommentsPart (0.5 points)
    # This verifies that the comments content itself has been removed.
    # FAILS on initial (3 comments exist) → PASSES on golden (no comments part or empty).
    try:
        comments_part_exists = False
        num_comments = 0
        comment_elements = []
        try:
            comments_part = doc.part.part_related_by(COMMENTS_REL)
            # If we get here, comments part exists — check if it has any w:comment elements
            comments_part_exists = len(doc.part.rels) > 0  # confirmed: part is loaded
            comments_root = comments_part._element
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            comment_elements = comments_root.findall('.//w:comment', ns)
            num_comments = len(comment_elements)
        except KeyError:
            # No comments relationship exists at all — fully deleted
            num_comments = 0

        if num_comments == 0:
            print(f"PASS: Component 1 — No comment elements found (CommentsPart present={comments_part_exists}, comments=0) (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Expected 0 comments, found {num_comments} comment element(s)")
            # Enumerate what comments exist for diagnostics
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            for c in comment_elements:
                cid = c.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', '?')
                # Get text
                texts = [t.text for t in c.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') if t.text]
                print(f"  Comment id={cid}: {''.join(texts)}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Comment range markers removed from document body (0.3 points)
    # commentRangeStart, commentRangeEnd, and commentReference elements must be absent.
    # FAILS on initial (3 of each) → PASSES on golden (0 of each).
    try:
        body_xml = etree.tostring(doc.element, encoding='unicode')
        cs_count = body_xml.count('w:commentRangeStart')
        ce_count = body_xml.count('w:commentRangeEnd')
        cr_count = body_xml.count('w:commentReference')
        total_markers = cs_count + ce_count + cr_count

        if total_markers == 0:
            print(f"PASS: Component 2 — No comment markers in document body (commentRangeStart=0, commentRangeEnd=0, commentReference=0) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Comment markers still present in document body: "
                  f"commentRangeStart={cs_count}, commentRangeEnd={ce_count}, commentReference={cr_count}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Document text is preserved — key section headings intact (0.2 points)
    # Verifies that deleting comments did NOT corrupt the document content.
    # The document should still have 26 paragraphs and all main section headings present.
    # This is a compound check: paragraphs exist AND text is intact.
    # FAILS on initial if coupled with Component 1 condition — we make it dependent:
    # we only award this if Component 1 passed (comments are gone) AND text is intact.
    # Actually: text preservation can independently pass in both states, so we gate it
    # on Component 1 having passed (score > 0) to ensure it only contributes when task is done.
    try:
        all_text = ' '.join(p.text for p in doc.paragraphs)
        headings_found = [h for h in EXPECTED_HEADINGS if h in all_text]
        num_paragraphs = len(doc.paragraphs)

        # Gate: only award if comments are actually deleted (Component 1 passed)
        if total_score >= 0.5 and len(headings_found) == len(EXPECTED_HEADINGS) and num_paragraphs >= 25:
            print(f"PASS: Component 3 — Document text preserved: {num_paragraphs} paragraphs, "
                  f"all {len(headings_found)}/{len(EXPECTED_HEADINGS)} section headings intact (0.2 pts)")
            total_score += 0.2
        elif total_score < 0.5:
            print(f"FAIL: Component 3 — Skipped (comments not yet deleted; gate not satisfied)")
        else:
            print(f"FAIL: Component 3 — Document text integrity issue: "
                  f"paragraphs={num_paragraphs}, headings found={len(headings_found)}/{len(EXPECTED_HEADINGS)}")
            missing = [h for h in EXPECTED_HEADINGS if h not in all_text]
            if missing:
                print(f"  Missing headings: {missing}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
