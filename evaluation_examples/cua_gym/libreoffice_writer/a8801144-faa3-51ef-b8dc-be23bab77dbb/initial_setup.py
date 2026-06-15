"""
Initial Setup: Document with a centered floating product image
Task ID: writer_obj_076
Domain: libreoffice_writer

Creates right_aligned_doc.docx at ~/Desktop/ with a 6cm x 4cm product image
currently positioned centered horizontally, and at 3cm from top of page.
The agent must move it to right-aligned and 5cm from top.
"""

import os
import shlex
import subprocess
import time
import io
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from PIL import Image, ImageDraw, ImageFont

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_076'
# Task context says file is at ~/Desktop/
OUTPUT = f'{WORKDIR}/Desktop/right_aligned_doc.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_product_image():
    """Create a realistic product image (6cm x 4cm)."""
    # 6cm x 4cm at 96 DPI: 6/2.54*96=226 x 4/2.54*96=151 pixels
    width_px = 226
    height_px = 151

    img = Image.new('RGB', (width_px, height_px), color=(240, 248, 255))
    draw = ImageDraw.Draw(img)

    # Draw a simple product box shape
    # Background gradient effect
    for y in range(height_px):
        shade = int(200 + (55 * y / height_px))
        draw.line([(0, y), (width_px, y)], fill=(shade, 220, shade))

    # Product box outline
    draw.rectangle([10, 10, width_px - 10, height_px - 10],
                   outline=(60, 120, 60), width=3)

    # Product label area
    draw.rectangle([20, 20, width_px - 20, 60],
                   fill=(60, 120, 60), outline=(40, 80, 40))

    # Simple text representation (drawn as rectangles for portability)
    # "PRODUCT" label - draw as white bars to simulate text
    bar_y = 32
    bar_positions = [25, 45, 65, 85, 105, 125, 145, 165]
    for bx in bar_positions:
        if bx + 14 < width_px - 20:
            draw.rectangle([bx, bar_y, bx + 12, bar_y + 10],
                           fill=(255, 255, 255))

    # Product details area
    for i in range(3):
        y_pos = 75 + i * 18
        draw.rectangle([25, y_pos, width_px - 25, y_pos + 8],
                       fill=(100, 160, 100))

    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes


def add_floating_image(doc, img_stream, width_cm, height_cm,
                       horiz_align='center', horiz_relative='margin',
                       vert_pos_cm=3.0, vert_relative='page'):
    """
    Add a floating image with anchor positioning.

    horiz_align: 'left', 'center', 'right'
    horiz_relative: 'margin', 'page', 'column', 'character', 'leftMargin', 'rightMargin'
    vert_pos_cm: vertical offset in cm from top of vert_relative
    vert_relative: 'page', 'margin', 'paragraph', 'line', 'topMargin', 'bottomMargin'
    """
    # Convert dimensions
    width_emu = int(width_cm * 360000)   # 1cm = 360000 EMU
    height_emu = int(height_cm * 360000)
    vert_emu = int(vert_pos_cm * 360000)

    # Add image to document relationships
    para = doc.add_paragraph()
    run = para.add_run()

    r_element = run._r

    img_stream.seek(0)

    # Use get_or_add_image to add image part and get relationship id
    rId, image = doc.part.get_or_add_image(img_stream)

    # Build anchor XML
    # distT/distB/distL/distR = distance from text (in EMU)
    dist = 114300  # ~0.32cm

    # Unique IDs for drawing
    draw_id = 1
    name = "ProductImage 1"

    # Build the full anchor XML string
    anchor_xml = f'''<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
    xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <wp:anchor distT="{dist}" distB="{dist}" distL="{dist}" distR="{dist}"
             simplePos="0" relativeHeight="251659264" behindDoc="0"
             locked="0" layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="{horiz_relative}">
      <wp:align>{horiz_align}</wp:align>
    </wp:positionH>
    <wp:positionV relativeFrom="{vert_relative}">
      <wp:posOffset>{vert_emu}</wp:posOffset>
    </wp:positionV>
    <wp:extent cx="{width_emu}" cy="{height_emu}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapSquare wrapText="bothSides"/>
    <wp:docPr id="{draw_id}" name="{name}" descr="Product Image"/>
    <wp:cNvGraphicFramePr>
      <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>
    </wp:cNvGraphicFramePr>
    <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
      <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
        <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
          <pic:nvPicPr>
            <pic:cNvPr id="0" name="{name}" descr="Product Image"/>
            <pic:cNvPicPr>
              <a:picLocks noChangeAspect="1" noChangeArrowheads="1"/>
            </pic:cNvPicPr>
          </pic:nvPicPr>
          <pic:blipFill>
            <a:blip r:embed="{rId}"/>
            <a:stretch>
              <a:fillRect/>
            </a:stretch>
          </pic:blipFill>
          <pic:spPr bwMode="auto">
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="{width_emu}" cy="{height_emu}"/>
            </a:xfrm>
            <a:prstGeom prst="rect">
              <a:avLst/>
            </a:prstGeom>
            <a:noFill/>
            <a:ln>
              <a:noFill/>
            </a:ln>
          </pic:spPr>
        </pic:pic>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>'''

    drawing_element = etree.fromstring(anchor_xml)
    run._r.append(drawing_element)

    return para


def create_initial():
    """Create initial document with a centered floating product image."""
    # Ensure Desktop directory exists
    desktop_dir = f'{WORKDIR}/Desktop'
    os.makedirs(desktop_dir, exist_ok=True)

    doc = Document()

    # Set page margins (standard A4 with 2.5cm margins)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Add document title
    title_para = doc.add_heading('EcoGreen Organic Products Catalog', level=1)

    # Add introduction paragraph
    intro = doc.add_paragraph(
        'Welcome to our 2025 product catalog. EcoGreen has been dedicated to providing '
        'sustainable, organic products to health-conscious consumers since 2010. '
        'Our commitment to quality and environmental responsibility sets us apart.'
    )

    # Add subtitle
    doc.add_heading('Featured Product: EcoGreen Harvest Blend', level=2)

    # Add product description
    desc1 = doc.add_paragraph(
        'The EcoGreen Harvest Blend is our flagship product, combining the finest organic '
        'ingredients sourced from certified farms across three continents. Each package '
        'contains a carefully curated selection of superfoods designed to support your '
        'daily nutritional needs.'
    )

    # Add empty paragraph where the floating image will appear
    img_anchor_para = doc.add_paragraph('')

    # Add more content
    doc.add_paragraph(
        'Key ingredients include organic spirulina, chia seeds, hemp protein, and '
        'a proprietary blend of 12 adaptogenic herbs. The product is vegan, gluten-free, '
        'and contains no artificial preservatives or additives.'
    )

    # Add product specifications section
    doc.add_heading('Product Specifications', level=2)

    # Add table with product details
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    specs = [
        ('Product Name', 'EcoGreen Harvest Blend 500g'),
        ('SKU', 'ECO-HB-500-2025'),
        ('Weight', '500g net / 520g gross'),
        ('Shelf Life', '24 months from manufacture date'),
        ('Certifications', 'USDA Organic, Non-GMO Project, Vegan Society'),
    ]
    for i, (label, value) in enumerate(specs):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        # Bold the label
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()

    # Add pricing section
    doc.add_heading('Pricing & Availability', level=2)
    doc.add_paragraph(
        'The EcoGreen Harvest Blend is available through authorized retailers and our '
        'online store. Wholesale pricing is available for orders of 50 units or more. '
        'Contact our sales team at sales@ecogreen.com for bulk order inquiries.'
    )

    # Add the floating image - centered, at 3cm from top (NOT the target position)
    img_stream = create_product_image()

    # Add to the document body - insert floating image paragraph after heading
    # Width: 6cm, Height: 4cm
    # Initial position: center horizontally, 3cm from top (NOT right, NOT 5cm)
    add_floating_image(
        doc, img_stream,
        width_cm=6.0, height_cm=4.0,
        horiz_align='center',    # Initial: center (task will change to right)
        horiz_relative='margin',
        vert_pos_cm=3.0,         # Initial: 3cm from top (task will change to 5cm)
        vert_relative='page'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
