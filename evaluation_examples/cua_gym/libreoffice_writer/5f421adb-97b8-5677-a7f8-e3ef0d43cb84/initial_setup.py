"""
Initial Setup: Create a Writer document with a 3x3 meeting agenda table
Task ID: writer_rd_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_041'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_cell_borders(cell, sz=4, color="000000", val="single"):
    """Set borders on a table cell. sz is in eighths of a point (4 = 0.5pt)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def create_initial():
    doc = Document()

    # Add a title
    heading = doc.add_heading('Q2 2025 Project Review Meeting', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph('Date: March 28, 2025 | Location: Conference Room B | Chair: Dr. Elena Vasquez')
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    intro.paragraph_format.space_after = Pt(12)

    # Create 3x3 table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    # Populate the table with meeting agenda content
    data = [
        ['Time Slot', 'Topic', 'Presenter'],
        ['9:00 - 9:45 AM', 'Budget Review & Resource Allocation for Q3', 'Marcus Chen, CFO'],
        ['9:45 - 10:30 AM', 'Product Roadmap Update: Mobile App v3.2', 'Action Items (see details)'],
    ]

    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            # Clear default paragraph and write
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(text)
            if i == 0:
                # Header row: bold
                run.bold = True
                run.font.size = Pt(11)
            else:
                run.font.size = Pt(10)
            # Set explicit 0.5pt solid black borders (sz=4 means 0.5pt in eighths)
            set_cell_borders(cell, sz=4, color="000000", val="single")

    # Add a closing paragraph
    doc.add_paragraph('')
    closing = doc.add_paragraph('Minutes prepared by: Sarah Thompson, Executive Assistant')
    closing.paragraph_format.space_before = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
