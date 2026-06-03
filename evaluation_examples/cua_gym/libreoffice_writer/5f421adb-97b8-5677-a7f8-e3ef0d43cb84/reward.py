"""
Reward Script: Nested table with dashed blue borders in last cell of 3x3 table
Task ID: writer_rd_041
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Nested table exists inside cell [2,2] of main table
  Component 2 (0.20): Nested table is 2x2
  Component 3 (0.20): Nested table borders are dashed style
  Component 4 (0.20): Nested table border color is blue (#0000CC)
  Component 5 (0.15): Nested table cells contain action item text (non-empty)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_041'


def persist_app_state(domain):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: main table must exist and be 3x3
    if len(doc.tables) < 1:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    main_table = doc.tables[0]
    if len(main_table.rows) < 3 or len(main_table.columns) < 3:
        print(f"FAIL: Main table is not at least 3x3 (got {len(main_table.rows)}x{len(main_table.columns)})")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    cell_2_2 = main_table.cell(2, 2)

    # Component 1: Nested table exists in cell [2,2] (0.25 points)
    nested_tbl = None
    try:
        nested_tbls = cell_2_2._element.findall('.//w:tbl', ns)
        if len(nested_tbls) > 0:
            nested_tbl = nested_tbls[0]
            print(f"PASS: Component 1 — Nested table found in cell [2,2] (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — No nested table found in cell [2,2]")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    if nested_tbl is None:
        # No nested table means remaining checks cannot pass
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Component 2: Nested table is 2x2 (0.20 points)
    try:
        nested_rows = nested_tbl.findall('.//w:tr', ns)
        num_rows = len(nested_rows)
        num_cols = 0
        if num_rows > 0:
            num_cols = len(nested_rows[0].findall('.//w:tc', ns))

        if num_rows == 2 and num_cols == 2:
            print(f"PASS: Component 2 — Nested table is 2x2 (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — Nested table is {num_rows}x{num_cols}, expected 2x2")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Nested table borders are dashed (0.20 points)
    try:
        tblPr = nested_tbl.find('w:tblPr', ns)
        tblBorders = tblPr.find('w:tblBorders', ns) if tblPr is not None else None

        dashed_count = 0
        border_names = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
        checked = 0

        if tblBorders is not None:
            for bname in border_names:
                border_el = tblBorders.find(f'w:{bname}', ns)
                if border_el is not None:
                    checked += 1
                    val = border_el.get(qn('w:val'))
                    if val and 'dash' in val.lower():
                        dashed_count += 1

        if checked > 0 and dashed_count >= 4:
            print(f"PASS: Component 3 — {dashed_count}/{checked} borders are dashed (0.20 pts)")
            total_score += 0.20
        else:
            # Fallback: check cell-level borders
            cell_dashed = 0
            cell_checked = 0
            for row_el in nested_rows:
                for tc in row_el.findall('.//w:tc', ns):
                    tcPr = tc.find('w:tcPr', ns)
                    if tcPr is not None:
                        tcBorders = tcPr.find('w:tcBorders', ns)
                        if tcBorders is not None:
                            for child in tcBorders:
                                cell_checked += 1
                                val = child.get(qn('w:val'))
                                if val and 'dash' in val.lower():
                                    cell_dashed += 1

            if cell_checked > 0 and cell_dashed >= 4:
                print(f"PASS: Component 3 — Cell-level: {cell_dashed}/{cell_checked} borders are dashed (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 — Not enough dashed borders (table-level: {dashed_count}/{checked}, cell-level: {cell_dashed}/{cell_checked})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Nested table border color is blue #0000CC (0.20 points)
    try:
        tblPr = nested_tbl.find('w:tblPr', ns)
        tblBorders = tblPr.find('w:tblBorders', ns) if tblPr is not None else None

        blue_count = 0
        color_checked = 0
        target_color = '0000CC'

        if tblBorders is not None:
            for bname in border_names:
                border_el = tblBorders.find(f'w:{bname}', ns)
                if border_el is not None:
                    color_checked += 1
                    color = border_el.get(qn('w:color'))
                    if color and color.upper().replace('#', '') == target_color.upper():
                        blue_count += 1

        if color_checked > 0 and blue_count >= 4:
            print(f"PASS: Component 4 — {blue_count}/{color_checked} borders are blue #0000CC (0.20 pts)")
            total_score += 0.20
        else:
            # Fallback: check cell-level borders
            cell_blue = 0
            cell_color_checked = 0
            for row_el in nested_rows:
                for tc in row_el.findall('.//w:tc', ns):
                    tcPr = tc.find('w:tcPr', ns)
                    if tcPr is not None:
                        tcBorders = tcPr.find('w:tcBorders', ns)
                        if tcBorders is not None:
                            for child in tcBorders:
                                cell_color_checked += 1
                                color = child.get(qn('w:color'))
                                if color and color.upper().replace('#', '') == target_color.upper():
                                    cell_blue += 1

            if cell_color_checked > 0 and cell_blue >= 4:
                print(f"PASS: Component 4 — Cell-level: {cell_blue}/{cell_color_checked} borders are blue (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 4 — Not enough blue borders (table-level: {blue_count}/{color_checked}, cell-level: {cell_blue}/{cell_color_checked})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Nested table cells contain text (action item details) (0.15 points)
    try:
        non_empty_cells = 0
        total_cells = 0
        for row_el in nested_rows:
            for tc in row_el.findall('.//w:tc', ns):
                total_cells += 1
                # Extract text from this cell
                cell_text = ''
                for p in tc.findall('.//w:p', ns):
                    for r in p.findall('.//w:r', ns):
                        for t in r.findall('.//w:t', ns):
                            cell_text += (t.text or '')
                if cell_text.strip():
                    non_empty_cells += 1

        if total_cells >= 4 and non_empty_cells >= 3:
            print(f"PASS: Component 5 — {non_empty_cells}/{total_cells} nested cells have content (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — Only {non_empty_cells}/{total_cells} nested cells have content (need at least 3 of 4)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
