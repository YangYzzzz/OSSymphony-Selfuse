"""
Initial Setup: Court declaration document without line numbering
Task ID: writer_legal_073
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_073'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard US Letter with legal margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # --- Case Caption ---
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cap.paragraph_format.space_after = Pt(6)
    run = cap.add_run('SUPERIOR COURT OF THE STATE OF CALIFORNIA')
    run.bold = True
    run.font.size = Pt(14)

    cap2 = doc.add_paragraph()
    cap2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cap2.paragraph_format.space_after = Pt(12)
    run2 = cap2.add_run('COUNTY OF LOS ANGELES')
    run2.bold = True
    run2.font.size = Pt(14)

    # Parties
    parties = doc.add_paragraph()
    parties.paragraph_format.space_after = Pt(6)
    r = parties.add_run('ELENA VASQUEZ and ROBERT VASQUEZ,')
    r.font.size = Pt(12)
    parties.add_run('\n')
    parties.add_run('                    Plaintiffs,').font.size = Pt(12)

    vs = doc.add_paragraph()
    vs.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    vs.paragraph_format.space_after = Pt(6)
    vs.add_run('v.').bold = True

    def_para = doc.add_paragraph()
    def_para.paragraph_format.space_after = Pt(6)
    def_para.add_run('PACIFIC RIDGE CONSTRUCTION, INC.,').font.size = Pt(12)
    def_para.add_run('\na California Corporation,')
    def_para.add_run('\n                    Defendant.')

    case_no = doc.add_paragraph()
    case_no.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    case_no.paragraph_format.space_after = Pt(12)
    r = case_no.add_run('Case No. BC-2025-04172')
    r.bold = True

    # --- Declaration Title ---
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(18)
    r = title.add_run('DECLARATION OF DR. MARGARET CHEN, Ph.D.')
    r.bold = True
    r.underline = True
    r.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run('IN SUPPORT OF PLAINTIFFS\' MOTION FOR SUMMARY JUDGMENT')
    r.bold = True
    r.font.size = Pt(12)

    # --- Declaration Body ---
    # Introductory paragraph
    intro = doc.add_paragraph()
    intro.paragraph_format.first_line_indent = Inches(0.5)
    intro.add_run(
        'I, Dr. Margaret Chen, Ph.D., declare under penalty of perjury under '
        'the laws of the State of California that the following is true and correct:'
    )

    # Numbered declaration paragraphs with substantial content
    declaration_paragraphs = [
        (
            'I am a licensed structural engineer in the State of California, '
            'License No. SE-78432, and have been continuously licensed since June 2008. '
            'I hold a Bachelor of Science in Civil Engineering from the University of '
            'California, Berkeley (2003), a Master of Science in Structural Engineering '
            'from Stanford University (2005), and a Doctor of Philosophy in Geotechnical '
            'Engineering from the California Institute of Technology (2008). I am a '
            'Fellow of the American Society of Civil Engineers and a member of the '
            'Structural Engineers Association of California.'
        ),
        (
            'I have been retained by counsel for Plaintiffs Elena Vasquez and Robert '
            'Vasquez to evaluate the structural integrity of the residential property '
            'located at 4728 Maple Ridge Drive, Pasadena, California 91107 (the '
            '"Property"). I have been compensated at my standard consulting rate of '
            '$450.00 per hour for my time in this matter. My opinions are based on my '
            'professional expertise and are not contingent upon the outcome of this '
            'litigation.'
        ),
        (
            'On September 14, 2025, I conducted an on-site inspection of the Property, '
            'which lasted approximately six hours. During my inspection, I was '
            'accompanied by my associate engineer, Mr. David Park, P.E. I took '
            'approximately 247 photographs documenting the conditions observed. I also '
            'collected twelve core samples from the foundation for laboratory analysis, '
            'the results of which are discussed below and attached hereto as Exhibit A.'
        ),
        (
            'The Property is a single-family residence originally constructed in 1987 by '
            'Pacific Ridge Construction, Inc. (the "Defendant"). The structure is a '
            'two-story wood-frame building with a post-tensioned concrete slab '
            'foundation, approximately 3,200 square feet of living space. According to '
            'the original building plans obtained from the City of Pasadena Department '
            'of Building and Safety (Permit No. BP-87-04291), the foundation was '
            'designed to support loads consistent with a residential occupancy '
            'classification, with minimum concrete compressive strength specified at '
            '3,000 pounds per square inch (psi).'
        ),
        (
            'My laboratory analysis of the twelve core samples revealed that the actual '
            'concrete compressive strength ranges from 1,840 psi to 2,350 psi, with a '
            'mean value of 2,067 psi. This represents a deficiency of approximately 31% '
            'below the minimum specified strength of 3,000 psi. The American Concrete '
            'Institute Standard ACI 318-19, Section 26.12.3.1, requires that the '
            'average of any three consecutive strength tests shall equal or exceed the '
            'specified compressive strength. The test results from the Property fail to '
            'meet this requirement by a substantial margin.'
        ),
        (
            'Furthermore, my inspection revealed significant cracking in the foundation '
            'slab. I identified seventeen distinct cracks exceeding 1/8 inch in width, '
            'with the largest crack measuring approximately 3/8 inch in width and '
            'extending 14 feet in length along the eastern perimeter of the foundation. '
            'The crack pattern is consistent with differential settlement caused by '
            'inadequate concrete strength and improper curing during the original '
            'construction. I observed evidence of prior cosmetic repairs to several '
            'cracks, consisting of surface-applied epoxy filler, which does not '
            'constitute a structural repair.'
        ),
        (
            'I also examined the post-tensioning system within the foundation slab. My '
            'inspection revealed that four of the twenty-eight post-tensioning tendons '
            'show evidence of corrosion and stress relaxation. Using a tendon stress '
            'meter, I measured residual stress levels in these four tendons ranging from '
            '62% to 71% of the original design stress. Industry standards, specifically '
            'the Post-Tensioning Institute\'s "Specification for Unbonded Single Strand '
            'Tendons" (PTI/ASBI M50.1-19), establish a minimum acceptable residual '
            'stress of 80% of the initial jacking stress. The loss of prestressing force '
            'in these tendons has contributed to the observed cracking and settlement.'
        ),
        (
            'I reviewed the soil investigation report prepared by Geotechnical Solutions, '
            'Inc. dated March 12, 1986, which was prepared in connection with the '
            'original construction. The report recommended a minimum foundation depth of '
            '36 inches below grade to reach competent bearing soil. However, my '
            'inspection of exposed foundation sections indicates that the actual '
            'foundation depth is approximately 24 inches below grade in several '
            'locations. This 12-inch shortfall in foundation depth is a significant '
            'deviation from the geotechnical engineer\'s recommendations and contributes '
            'to the observed differential settlement.'
        ),
        (
            'To assess the extent of structural damage to the superstructure, I '
            'conducted a level survey of the first and second floors using a Leica '
            'NA730 automatic level with a stated accuracy of +/- 0.7mm per kilometer. '
            'The survey revealed maximum floor elevation differences of 1.25 inches '
            'across the first floor and 1.75 inches across the second floor. These '
            'measurements exceed the tolerance of L/360, where L is the span length, '
            'as specified by the International Building Code Section 1604.3.1 for '
            'floor members under live load deflection limits.'
        ),
        (
            'The differential settlement has caused visible damage to interior finishes '
            'throughout the Property. I documented the following: (a) diagonal cracking '
            'in drywall at twenty-three locations, primarily at window and door corners; '
            '(b) separation of crown molding from ceiling surfaces in the living room '
            'and master bedroom, with gaps measuring up to 1/2 inch; (c) misalignment '
            'of six interior doors such that they no longer close properly; (d) cracking '
            'of ceramic tile flooring in the kitchen and both bathrooms; and (e) visible '
            'separation at the junction of interior partition walls and exterior bearing '
            'walls at four locations.'
        ),
        (
            'I also observed water intrusion evidence at three foundation crack locations. '
            'Moisture meter readings taken at these locations on September 14, 2025, '
            'showed moisture content levels ranging from 28% to 34% in the adjacent wood '
            'framing members. The Forest Products Laboratory establishes that wood '
            'moisture content above 20% creates conditions favorable for fungal decay. '
            'My visual inspection confirmed the presence of white rot fungus on floor '
            'joists at two of these locations, which has compromised the structural '
            'capacity of the affected members.'
        ),
        (
            'Based on my analysis, I have prepared a remediation plan and cost estimate '
            'for the necessary structural repairs. The remediation would require: '
            '(a) underpinning of the foundation to proper depth using helical piers at '
            'approximately 32 locations, estimated cost $128,000; (b) carbon fiber '
            'reinforcement of the foundation slab to restore structural capacity, '
            'estimated cost $67,500; (c) re-stressing or replacement of the four '
            'deficient post-tensioning tendons, estimated cost $42,000; (d) replacement '
            'of fungal-damaged floor joists and associated framing, estimated cost '
            '$38,500; and (e) repair and restoration of interior finishes, estimated '
            'cost $94,000. The total estimated cost of remediation is $370,000.'
        ),
        (
            'It is my professional opinion, held to a reasonable degree of engineering '
            'certainty, that the structural deficiencies observed at the Property are '
            'the direct result of substandard construction practices during the original '
            'construction in 1987. Specifically, the use of concrete with compressive '
            'strength 31% below specification, the failure to achieve the recommended '
            'foundation depth, and deficiencies in the post-tensioning system are all '
            'attributable to the original construction by Pacific Ridge Construction, '
            'Inc. These deficiencies were not the result of normal wear and aging, '
            'subsequent modifications by the Plaintiffs, or natural forces such as '
            'earthquake or flooding.'
        ),
        (
            'I further opine that the structural deficiencies described above render the '
            'Property unsafe for continued residential occupancy without the remediation '
            'described in Paragraph 12. The risk of continued deterioration is '
            'substantial, as the ongoing differential settlement will progressively '
            'worsen the cracking in both the foundation and superstructure. Additionally, '
            'the water intrusion and fungal decay, if left unaddressed, will continue to '
            'compromise the structural integrity of floor framing members, creating a '
            'risk of localized floor collapse.'
        ),
        (
            'In reaching my conclusions, I have relied upon the following materials: '
            '(a) my personal observations during the September 14, 2025 site inspection; '
            '(b) the original building plans and specifications (Permit No. BP-87-04291); '
            '(c) the Geotechnical Solutions soil investigation report dated March 12, '
            '1986; (d) laboratory test results from Pacific Testing Laboratories for the '
            'twelve concrete core samples; (e) the tendon stress measurements taken on '
            'site; (f) my level survey data; (g) the applicable provisions of ACI 318-19, '
            'the International Building Code 2021 Edition, and PTI/ASBI M50.1-19; and '
            '(h) my education, training, and over seventeen years of experience in '
            'structural engineering.'
        ),
        (
            'I have reviewed the deposition testimony of Mr. Frank Delgado, the project '
            'superintendent employed by Pacific Ridge Construction during the '
            'construction of the Property. Mr. Delgado testified that "we had some '
            'problems with the concrete delivery on the day of the pour" and that "the '
            'batch plant sent a mix that seemed a little thin, but we went ahead with it '
            'because we were already behind schedule." This testimony is consistent with '
            'my finding that the concrete used in the foundation was below the specified '
            'compressive strength and supports the conclusion that the Defendant was '
            'aware of potential quality issues at the time of construction.'
        ),
        (
            'I am available to testify at trial and to respond to any questions '
            'regarding my findings, methodology, and opinions set forth in this '
            'declaration. My curriculum vitae, which details my qualifications, '
            'publications, and prior expert witness experience, is attached hereto as '
            'Exhibit B. In the past four years, I have been retained as an expert '
            'witness in approximately twenty-two cases involving residential construction '
            'defects, and I have testified at trial or deposition in fourteen of those '
            'matters.'
        ),
    ]

    for i, text in enumerate(declaration_paragraphs, 1):
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Inches(0.5)
        run_num = para.add_run(f'{i}.    ')
        run_num.bold = True
        para.add_run(text)

    # Signature block
    doc.add_paragraph()  # blank line

    sig_intro = doc.add_paragraph()
    sig_intro.paragraph_format.first_line_indent = Inches(0.5)
    sig_intro.add_run(
        'I declare under penalty of perjury under the laws of the State of California '
        'that the foregoing is true and correct.'
    )

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.paragraph_format.first_line_indent = Inches(0.5)
    date_para.add_run('Executed on October 7, 2025, at Pasadena, California.')

    doc.add_paragraph()
    doc.add_paragraph()

    sig_line = doc.add_paragraph()
    sig_line.paragraph_format.left_indent = Inches(3.5)
    sig_line.add_run('_________________________________')

    name_line = doc.add_paragraph()
    name_line.paragraph_format.left_indent = Inches(3.5)
    r = name_line.add_run('Dr. Margaret Chen, Ph.D., S.E.')
    r.bold = True

    title_line = doc.add_paragraph()
    title_line.paragraph_format.left_indent = Inches(3.5)
    title_line.add_run('California Licensed Structural Engineer')

    lic_line = doc.add_paragraph()
    lic_line.paragraph_format.left_indent = Inches(3.5)
    lic_line.add_run('License No. SE-78432')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
