"""
Reward Script: ADA Workplace Accommodation Request Form
Task ID: writer_hr_087
Domain: libreoffice_writer
Scoring:
  Component 1: Employee Information section with labeled fields (0.20)
  Component 2: Confidentiality notice in medical/disability section (0.15)
  Component 3: Essential job functions table with 5 rows and checkbox columns (0.25)
  Component 4: Three accommodation request blocks with Description/Duration/Cost (0.20)
  Component 5: Approval workflow table with Manager/HR/Legal rows and checkboxes (0.20)
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_087'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_text = "\n".join(p.text for p in doc.paragraphs).lower()
    headings = [(p.style.name, p.text) for p in doc.paragraphs if p.style and 'Heading' in p.style.name]

    # ======================================================================
    # Component 1: Employee Information section with labeled fields (0.20 pts)
    # The initial file has NO employee info section — only title + disclaimer.
    # ======================================================================
    try:
        # Check for an Employee Information heading
        has_emp_heading = any("employee information" in h[1].lower() for h in headings)
        # Check for at least 4 of these employee info field labels
        emp_fields = ["full name", "employee id", "department", "position", "date of request",
                      "supervisor", "work location", "phone", "email"]
        found_fields = sum(1 for f in emp_fields if f in all_text)

        if has_emp_heading and found_fields >= 4:
            print(f"PASS: Component 1 — Employee Info heading found, {found_fields}/9 fields present (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — heading={has_emp_heading}, fields={found_fields}/9 (need heading + >=4 fields)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ======================================================================
    # Component 2: Confidentiality notice in medical/disability section (0.15 pts)
    # The initial file has NO medical section or confidentiality notice.
    # ======================================================================
    try:
        # Check for a Disability/Medical heading
        has_medical_heading = any(
            "disability" in h[1].lower() or "medical" in h[1].lower()
            for h in headings
        )
        # Check for confidentiality notice text
        has_confidentiality = "confidentiality" in all_text and (
            "hipaa" in all_text or "protected" in all_text or "shared" in all_text
        )

        if has_medical_heading and has_confidentiality:
            print(f"PASS: Component 2 — Medical section heading + confidentiality notice found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 — medical_heading={has_medical_heading}, confidentiality={has_confidentiality}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ======================================================================
    # Component 3: Essential job functions table with 5 rows and checkbox cols (0.25 pts)
    # The initial file has NO tables at all.
    # ======================================================================
    try:
        functions_table = None
        for table in doc.tables:
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            header_text = " ".join(header_cells)
            if ("can perform" in header_text and "cannot perform" in header_text) or \
               ("essential" in header_text and "function" in header_text):
                functions_table = table
                break
            # Also check if column headers match
            if len(table.columns) >= 3:
                if any("can" in c and "perform" in c for c in header_cells):
                    functions_table = table
                    break

        if functions_table is None:
            print("FAIL: Component 3 — No essential job functions table found")
        else:
            num_data_rows = len(functions_table.rows) - 1  # exclude header
            # Check for checkbox characters in table
            checkbox_chars = {"☐", "☑", "☒", "□", "■", "✓", "✗"}
            has_checkboxes = any(
                any(ch in cell.text for ch in checkbox_chars)
                for row in functions_table.rows[1:]
                for cell in row.cells
            )

            sub_score = 0.0
            # 5 essential functions listed (0.15)
            if num_data_rows >= 5:
                sub_score += 0.15
                print(f"  PASS: 3a — {num_data_rows} function rows (need >=5)")
            else:
                print(f"  FAIL: 3a — only {num_data_rows} function rows (need >=5)")

            # Checkboxes present (0.10)
            if has_checkboxes:
                sub_score += 0.10
                print(f"  PASS: 3b — Checkbox characters found in table")
            else:
                print(f"  FAIL: 3b — No checkbox characters found in functions table")

            if sub_score > 0:
                print(f"PASS: Component 3 — Essential functions table ({sub_score:.2f} pts)")
                total_score += sub_score
            else:
                print(f"FAIL: Component 3 — Table found but no sub-checks passed")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ======================================================================
    # Component 4: Three accommodation request blocks (0.20 pts)
    # Each block should have Description, Expected Duration, Cost Estimate.
    # The initial file has NONE of these.
    # ======================================================================
    try:
        # Count accommodation request sub-headings or numbered blocks
        accom_headings = [h for h in headings if "accommodation" in h[1].lower() and (
            "#" in h[1] or "request" in h[1].lower()
        )]

        # Also search for accommodation blocks by text patterns
        accom_count_by_text = 0
        for p in doc.paragraphs:
            txt = p.text.lower()
            if re.search(r'accommodation\s+(request\s+)?#?\s*\d', txt):
                accom_count_by_text += 1

        num_accom = max(len(accom_headings), accom_count_by_text)

        # Check for required sub-fields in the accommodation section
        has_description = "description" in all_text and "accommodation" in all_text
        has_duration = "duration" in all_text or "expected duration" in all_text
        has_cost = "cost" in all_text or "estimate" in all_text

        sub_score = 0.0
        if num_accom >= 3:
            sub_score += 0.10
            print(f"  PASS: 4a — {num_accom} accommodation request blocks found")
        else:
            print(f"  FAIL: 4a — only {num_accom} accommodation blocks (need >=3)")

        fields_present = sum([has_description, has_duration, has_cost])
        if fields_present >= 3:
            sub_score += 0.10
            print(f"  PASS: 4b — All 3 sub-fields found (Description, Duration, Cost)")
        elif fields_present >= 2:
            sub_score += 0.05
            print(f"  FAIL: 4b — Only {fields_present}/3 sub-fields (Description={has_description}, Duration={has_duration}, Cost={has_cost})")
        else:
            print(f"  FAIL: 4b — Only {fields_present}/3 sub-fields")

        if sub_score > 0:
            print(f"PASS: Component 4 — Accommodation requests ({sub_score:.2f} pts)")
            total_score += sub_score
        else:
            print(f"FAIL: Component 4 — No accommodation request blocks found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ======================================================================
    # Component 5: Approval workflow table with Manager/HR/Legal (0.20 pts)
    # The initial file has NO approval workflow section or table.
    # ======================================================================
    try:
        # Check for approval workflow heading
        has_approval_heading = any("approval" in h[1].lower() or "workflow" in h[1].lower() for h in headings)

        # Find the approval table — should have Manager, HR, Legal rows
        approval_table = None
        for table in doc.tables:
            table_text = " ".join(cell.text.strip().lower() for row in table.rows for cell in row.cells)
            if "manager" in table_text and ("hr" in table_text or "human resources" in table_text) and "legal" in table_text:
                approval_table = table
                break

        sub_score = 0.0

        if has_approval_heading:
            sub_score += 0.05
            print(f"  PASS: 5a — Approval workflow heading found")
        else:
            print(f"  FAIL: 5a — No approval workflow heading found")

        if approval_table is not None:
            # Check for checkbox characters in approval table
            checkbox_chars = {"☐", "☑", "☒", "□", "■", "✓", "✗"}
            has_checkboxes = False
            roles_found = set()
            for row in approval_table.rows:
                row_text = " ".join(cell.text.strip().lower() for cell in row.cells)
                if "manager" in row_text:
                    roles_found.add("manager")
                if "hr" in row_text or "human resources" in row_text:
                    roles_found.add("hr")
                if "legal" in row_text:
                    roles_found.add("legal")
                if not has_checkboxes and any(ch in row_text for ch in checkbox_chars):
                    has_checkboxes = not False  # checkbox detected via character scan

            if len(roles_found) >= 3:
                sub_score += 0.10
                print(f"  PASS: 5b — All 3 reviewer roles found: {roles_found}")
            else:
                print(f"  FAIL: 5b — Only {len(roles_found)}/3 roles found: {roles_found}")

            if has_checkboxes:
                sub_score += 0.05
                print(f"  PASS: 5c — Checkboxes found in approval table")
            else:
                print(f"  FAIL: 5c — No checkboxes in approval table")
        else:
            print(f"  FAIL: 5b/5c — No approval workflow table found")

        if sub_score > 0:
            print(f"PASS: Component 5 — Approval workflow ({sub_score:.2f} pts)")
            total_score += sub_score
        else:
            print(f"FAIL: Component 5 — Approval workflow section not found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
