"""
Initial Setup: ADA Workplace Accommodation Request Form - minimal state
Task ID: writer_hr_087
Domain: libreoffice_writer

Creates document with only a title and legal disclaimer paragraph.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_087'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title ---
    title = doc.add_heading('ADA Workplace Accommodation Request Form', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Legal Disclaimer ---
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = disclaimer.add_run(
        'LEGAL DISCLAIMER: This form is used to initiate a request for reasonable '
        'accommodation under the Americans with Disabilities Act (ADA) of 1990, as '
        'amended, and Section 504 of the Rehabilitation Act of 1973. All information '
        'provided on this form is strictly confidential and will be maintained in a '
        'separate medical file in accordance with 29 CFR 1630.14(c). Disclosure of '
        'medical information will be limited to those with a legitimate need to know. '
        'Completion of this form does not guarantee that an accommodation will be '
        'granted. The employer will engage in an interactive process with the employee '
        'to identify effective accommodations.'
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
