"""
Initial Setup: Format WHEREAS clauses in legal contract
Task ID: writer_legal_070
Domain: libreoffice_writer

Creates a legal contract with 6 WHEREAS clauses in regular formatting.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_070'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('ASSET PURCHASE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Date and Parties ---
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Effective Date: March 15, 2025')
    run.font.size = Pt(11)

    doc.add_paragraph('')  # spacer

    parties = doc.add_paragraph()
    parties_run = parties.add_run(
        'This Asset Purchase Agreement (the "Agreement") is entered into as of '
        'the date first written above, by and between Meridian Technology Solutions, Inc., '
        'a Delaware corporation with its principal office at 4200 Lakewood Boulevard, Suite 300, '
        'Long Beach, CA 90808 (the "Buyer"), and Pinnacle Data Systems, LLC, a California '
        'limited liability company with its principal office at 1750 Montgomery Street, '
        'San Francisco, CA 94111 (the "Seller").'
    )
    parties_run.font.size = Pt(11)
    parties_run.font.name = 'Times New Roman'

    doc.add_paragraph('')  # spacer

    recitals_heading = doc.add_heading('RECITALS', level=1)
    recitals_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- 6 WHEREAS clauses (all regular formatting, NO bold, NO small caps) ---
    whereas_clauses = [
        (
            'WHEREAS',
            ', the Seller is engaged in the business of providing enterprise data management, '
            'cloud infrastructure services, and cybersecurity consulting to commercial and '
            'government clients throughout the Western United States (the "Business");'
        ),
        (
            'WHEREAS',
            ', the Seller desires to sell, transfer, and assign to the Buyer, and the Buyer '
            'desires to purchase and acquire from the Seller, substantially all of the assets '
            'used in or related to the Business, subject to the terms and conditions set forth '
            'in this Agreement;'
        ),
        (
            'WHEREAS',
            ', the Buyer has conducted a thorough due diligence investigation of the Business, '
            'including a review of the Seller\'s financial statements for fiscal years 2022, '
            '2023, and 2024, and is satisfied with the results of such investigation;'
        ),
        (
            'WHEREAS',
            ', the Board of Directors of each party has determined that the transactions '
            'contemplated by this Agreement are in the best interests of their respective '
            'companies and stakeholders, and has approved and authorized the execution and '
            'delivery of this Agreement;'
        ),
        (
            'WHEREAS',
            ', the parties acknowledge that the aggregate purchase price of Twelve Million '
            'Five Hundred Thousand Dollars ($12,500,000.00) represents the fair market value '
            'of the Acquired Assets as determined by independent valuation conducted by '
            'Hartfield & Associates, LLP;'
        ),
        (
            'WHEREAS',
            ', the Seller has obtained all necessary consents, approvals, and waivers from '
            'third parties, including key customers, vendors, and licensors, required for the '
            'valid transfer of the Acquired Assets to the Buyer pursuant to this Agreement;'
        ),
    ]

    for whereas_word, clause_text in whereas_clauses:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        # "WHEREAS" in regular formatting (same as rest of text)
        run_whereas = para.add_run(whereas_word)
        run_whereas.font.size = Pt(11)
        run_whereas.font.name = 'Times New Roman'
        run_whereas.bold = False
        # Rest of clause text
        run_text = para.add_run(clause_text)
        run_text.font.size = Pt(11)
        run_text.font.name = 'Times New Roman'

    doc.add_paragraph('')  # spacer

    # --- Agreement section ---
    agreement_heading = doc.add_heading('AGREEMENT', level=1)
    agreement_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    now_para = doc.add_paragraph()
    now_run = now_para.add_run(
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements hereinafter '
        'set forth and for other good and valuable consideration, the receipt and sufficiency '
        'of which are hereby acknowledged, the parties agree as follows:'
    )
    now_run.font.size = Pt(11)
    now_run.font.name = 'Times New Roman'

    # Article 1
    art1 = doc.add_heading('Article 1. Definitions', level=2)
    p1 = doc.add_paragraph()
    r1 = p1.add_run(
        '1.1 "Acquired Assets" means all right, title, and interest in and to all assets '
        'owned, held, or used by the Seller in connection with the Business, including but '
        'not limited to tangible personal property, intellectual property, contract rights, '
        'customer lists, and goodwill.'
    )
    r1.font.size = Pt(11)
    r1.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        '1.2 "Closing Date" means the date on which the transactions contemplated by this '
        'Agreement are consummated, which shall be no later than April 30, 2025, or such '
        'other date as the parties may mutually agree upon in writing.'
    )
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

    # Article 2
    art2 = doc.add_heading('Article 2. Purchase and Sale of Assets', level=2)
    p3 = doc.add_paragraph()
    r3 = p3.add_run(
        '2.1 Purchase and Sale. Subject to the terms and conditions of this Agreement, at '
        'the Closing, the Seller shall sell, assign, transfer, convey, and deliver to the '
        'Buyer, and the Buyer shall purchase, acquire, and accept from the Seller, all of '
        'the Acquired Assets, free and clear of all liens, encumbrances, and claims.'
    )
    r3.font.size = Pt(11)
    r3.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
