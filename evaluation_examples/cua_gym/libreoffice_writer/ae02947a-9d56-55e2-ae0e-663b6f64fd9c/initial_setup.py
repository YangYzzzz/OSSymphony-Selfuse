"""
Initial Setup: Financial Analysis document with chart image, no caption.
Task ID: writer_frd_076
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Generate a simple bar chart image using PIL
from PIL import Image, ImageDraw, ImageFont

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_076'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
CHART_IMG = f'{WORKDIR}/annual_revenue_chart.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_chart_image():
    """Create a realistic-looking bar chart image for the document."""
    width, height = 640, 400
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)

    # Title
    try:
        title_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 18)
        label_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 12)
        axis_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 11)
    except IOError:
        title_font = ImageFont.load_default()
        label_font = ImageFont.load_default()
        axis_font = ImageFont.load_default()

    draw.text((width // 2 - 120, 15), "Annual Revenue Growth", fill='black', font=title_font)

    # Chart area
    chart_left, chart_top = 80, 60
    chart_right, chart_bottom = 600, 340
    draw.rectangle([chart_left, chart_top, chart_right, chart_bottom], outline='gray')

    # Data: revenue in millions for 2019-2024
    years = ['2019', '2020', '2021', '2022', '2023', '2024']
    revenues = [12.4, 14.8, 18.2, 22.5, 27.1, 33.6]
    max_rev = 40.0
    bar_width = 55
    gap = 30
    start_x = chart_left + 25

    colors = ['#2E5090', '#3A6BAF', '#4682B4', '#5A9BD5', '#70B8E8', '#87CEEB']

    for i, (year, rev) in enumerate(zip(years, revenues)):
        bar_height = int((rev / max_rev) * (chart_bottom - chart_top - 20))
        x1 = start_x + i * (bar_width + gap)
        y1 = chart_bottom - bar_height
        x2 = x1 + bar_width
        y2 = chart_bottom

        draw.rectangle([x1, y1, x2, y2], fill=colors[i], outline='#1A3660')
        draw.text((x1 + 5, y1 - 16), f"${rev}M", fill='black', font=axis_font)
        draw.text((x1 + 10, chart_bottom + 5), year, fill='black', font=label_font)

    # Y-axis labels
    for val in [0, 10, 20, 30, 40]:
        y = chart_bottom - int((val / max_rev) * (chart_bottom - chart_top - 20))
        draw.text((chart_left - 30, y - 6), f"${val}M", fill='gray', font=axis_font)
        draw.line([(chart_left, y), (chart_right, y)], fill='#DDDDDD', width=1)

    # Y-axis title
    draw.text((10, height // 2 - 40), "Revenue\n(USD M)", fill='black', font=axis_font)

    img.save(CHART_IMG)
    print(f"Chart image created: {CHART_IMG}")


def create_initial():
    create_chart_image()

    doc = Document()

    # --- Title ---
    title = doc.add_heading('Meridian Technologies Inc.', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_heading('Annual Financial Analysis Report', level=1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'This report presents a comprehensive analysis of Meridian Technologies Inc.\'s '
        'financial performance over the period 2019 to 2024. The company has demonstrated '
        'consistent revenue growth, driven primarily by expansion in cloud services and '
        'enterprise software licensing. Operating margins have improved from 18.3% to 24.7% '
        'during this period, reflecting successful cost optimization initiatives and '
        'economies of scale.'
    )

    # --- Revenue Performance ---
    doc.add_heading('Revenue Performance', level=2)
    doc.add_paragraph(
        'Meridian Technologies achieved total revenue of $33.6 million in fiscal year 2024, '
        'representing a 24.0% year-over-year increase from $27.1 million in 2023. This marks '
        'the sixth consecutive year of double-digit growth. The cloud services division '
        'contributed $19.2 million (57.1% of total revenue), while enterprise licensing '
        'accounted for $10.8 million (32.1%), and professional services generated $3.6 million '
        '(10.7%).'
    )
    doc.add_paragraph(
        'The following chart illustrates the company\'s revenue trajectory over the past six '
        'fiscal years:'
    )

    # --- Insert chart image (NO CAPTION) ---
    doc.add_picture(CHART_IMG, width=Inches(5.5))
    # Center the image paragraph
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Key Drivers ---
    doc.add_heading('Key Growth Drivers', level=2)
    doc.add_paragraph(
        'Several strategic initiatives have contributed to the sustained revenue growth:'
    )
    doc.add_paragraph(
        'Cloud Platform Expansion: The launch of Meridian Cloud Suite 3.0 in Q2 2023 '
        'attracted 340 new enterprise clients, representing a 28% increase in the cloud '
        'customer base.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Geographic Diversification: International revenue grew by 31% year-over-year, '
        'with the Asia-Pacific region contributing $8.4 million, up from $5.9 million in 2023.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Strategic Acquisitions: The acquisition of DataStream Analytics in March 2024 '
        'added $2.1 million in recurring revenue and expanded the company\'s data analytics '
        'capabilities.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Customer Retention: Net revenue retention rate improved to 118%, indicating strong '
        'upsell and cross-sell performance within the existing customer base.',
        style='List Bullet'
    )

    # --- Profitability Analysis ---
    doc.add_heading('Profitability Analysis', level=2)
    doc.add_paragraph(
        'Gross profit margin expanded to 71.2% in 2024, compared to 68.5% in 2023. This '
        'improvement was primarily driven by the shift toward higher-margin cloud subscription '
        'revenue and reduced infrastructure costs following the migration to a multi-tenant '
        'architecture. EBITDA reached $9.8 million, representing a 29.2% EBITDA margin, '
        'compared to $7.3 million (26.9% margin) in the prior year.'
    )

    # --- Outlook ---
    doc.add_heading('Forward Outlook', level=2)
    doc.add_paragraph(
        'Management projects revenue of $41 to $43 million for fiscal year 2025, representing '
        'approximately 22-28% growth. Key investments planned include the expansion of the '
        'European data center network, development of AI-powered analytics features, and '
        'continued expansion of the partner ecosystem. The company anticipates maintaining '
        'operating margins above 25% while continuing to invest in growth initiatives.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
