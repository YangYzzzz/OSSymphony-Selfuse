"""
Reward Script: Add caption 'Figure 1: Annual Revenue Growth' below chart image
Task ID: writer_frd_076
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.30): Caption paragraph with correct text exists
  - Component 2 (0.25): Caption uses 'Caption' paragraph style
  - Component 3 (0.25): Caption contains SEQ Figure field (automatic numbering)
  - Component 4 (0.20): Caption is positioned immediately after the image paragraph
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_076'
EXPECTED_CAPTION_TEXT = 'Figure 1: Annual Revenue Growth'

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': WNS}


def find_caption_paragraph(doc):
    """Find the paragraph that contains the caption text."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == EXPECTED_CAPTION_TEXT:
            return i, para
    # Also check with flexible matching (case insensitive, whitespace tolerant)
    for i, para in enumerate(doc.paragraphs):
        text = ' '.join(para.text.strip().split())
        if text.lower() == EXPECTED_CAPTION_TEXT.lower():
            return i, para
    return None, None


def find_image_paragraph_index(doc):
    """Find the index of the paragraph containing the chart image."""
    for i, para in enumerate(doc.paragraphs):
        xml = para._element.xml
        if 'blip' in xml or 'graphicData' in xml:
            return i
    return None


def has_seq_figure_field(para):
    """Check if a paragraph contains a SEQ Figure field (automatic numbering)."""
    instr_texts = para._element.findall('.//w:instrText', NS)
    for instr in instr_texts:
        if instr.text and 'SEQ' in instr.text and 'Figure' in instr.text:
            return True
    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Caption paragraph with correct text exists (0.30 points)
    try:
        cap_idx, cap_para = find_caption_paragraph(doc)
        if cap_para is not None:
            print(f"PASS: Component 1 — Caption paragraph found at P{cap_idx}: '{cap_para.text.strip()}' (0.30 pts)")
            total_score += 0.30
        else:
            # Check if any paragraph contains partial caption text
            found_any = False
            for i, p in enumerate(doc.paragraphs):
                if 'Annual Revenue Growth' in p.text:
                    print(f"FAIL: Component 1 — Found partial match at P{i}: '{p.text.strip()}' but expected exact: '{EXPECTED_CAPTION_TEXT}'")
                    found_any = True
                    break
            if not found_any:
                print(f"FAIL: Component 1 — No paragraph contains caption text '{EXPECTED_CAPTION_TEXT}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Caption uses 'Caption' paragraph style (0.25 points)
    try:
        if cap_para is not None:
            style_name = cap_para.style.name if cap_para.style else 'None'
            # Accept 'Caption' or style names containing 'caption' (case insensitive)
            if style_name == 'Caption' or 'caption' in style_name.lower():
                print(f"PASS: Component 2 — Caption style is '{style_name}' (0.25 pts)")
                total_score += 0.25
            else:
                # Also check the raw pStyle value in XML
                pPr = cap_para._element.find('.//w:pPr/w:pStyle', NS)
                raw_style = pPr.get(qn('w:val')) if pPr is not None else None
                if raw_style and 'caption' in raw_style.lower():
                    print(f"PASS: Component 2 — Caption raw style is '{raw_style}' (0.25 pts)")
                    total_score += 0.25
                else:
                    print(f"FAIL: Component 2 — Caption style is '{style_name}' (raw: {raw_style}), expected 'Caption'")
        else:
            print(f"FAIL: Component 2 — No caption paragraph found to check style")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Caption contains SEQ Figure field for automatic numbering (0.25 points)
    try:
        if cap_para is not None:
            if has_seq_figure_field(cap_para):
                print(f"PASS: Component 3 — SEQ Figure field found in caption (0.25 pts)")
                total_score += 0.25
            else:
                # Check for any field code at all
                fld_chars = cap_para._element.findall('.//w:fldChar', NS)
                if fld_chars:
                    instr_texts = cap_para._element.findall('.//w:instrText', NS)
                    instr_content = [it.text for it in instr_texts]
                    print(f"FAIL: Component 3 — Found field codes but not SEQ Figure. instrTexts: {instr_content}")
                else:
                    print(f"FAIL: Component 3 — No field codes found in caption (no automatic numbering)")
        else:
            print(f"FAIL: Component 3 — No caption paragraph found to check SEQ field")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Caption is positioned immediately after the image paragraph (0.20 points)
    try:
        img_idx = find_image_paragraph_index(doc)
        if img_idx is not None and cap_idx is not None:
            if cap_idx == img_idx + 1:
                print(f"PASS: Component 4 — Caption at P{cap_idx} is right after image at P{img_idx} (0.20 pts)")
                total_score += 0.20
            elif cap_idx == img_idx + 2 and doc.paragraphs[img_idx + 1].text.strip() == '':
                # Allow one blank paragraph between image and caption
                print(f"PASS: Component 4 — Caption at P{cap_idx} is after image at P{img_idx} (blank para between) (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 4 — Caption at P{cap_idx}, image at P{img_idx} (expected caption right after image)")
        elif img_idx is None:
            print(f"FAIL: Component 4 — No image paragraph found in document")
        else:
            print(f"FAIL: Component 4 — No caption paragraph found to check position")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
