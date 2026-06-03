"""
Initial Setup: Multi-level list where all levels use the same default round bullet
Task ID: writer_list_049
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

WORKDIR = '/home/user'
TASK_ID = 'topic_outline'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_numbering_with_same_bullet(doc):
    """
    Create an abstract numbering definition in numbering.xml where
    all three levels (0, 1, 2) use the same default round bullet U+2022.
    Returns the numId to reference from paragraphs.
    """
    # Access or create numbering part
    try:
        numbering_part = doc.part.numbering_part
        numbering = numbering_part._element
    except AttributeError:
        # Create numbering part from scratch
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from lxml import etree

        numbering_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:numbering xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"'
            ' xmlns:mo="http://schemas.microsoft.com/office/mac/office/2008/main"'
            ' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"'
            ' xmlns:mv="urn:schemas-microsoft-com:mac:vml"'
            ' xmlns:o="urn:schemas-microsoft-com:office:office"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
            ' xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
            ' xmlns:v="urn:schemas-microsoft-com:vml"'
            ' xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"'
            ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
            ' xmlns:w10="urn:schemas-microsoft-com:office:word"'
            ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"'
            ' xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"'
            ' xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"'
            ' xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"'
            ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
            ' mc:Ignorable="w14 wp14">'
            '</w:numbering>'
        )
        numbering = etree.fromstring(numbering_xml.encode('utf-8'))

        from docx.opc.part import Part
        numbering_part = Part(
            PackURI('/word/numbering.xml'),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml',
            etree.tostring(numbering, xml_declaration=True, encoding='UTF-8', standalone=True),
            doc.part.package
        )
        doc.part.relate_to(
            numbering_part,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering'
        )
        numbering = numbering_part._element

    # Build abstractNum element with 3 levels, all using U+2022 bullet
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def make_el(tag, attribs=None):
        el = OxmlElement(tag)
        if attribs:
            for k, v in attribs.items():
                el.set(qn(k) if ':' not in k else k, v)
        return el

    # Determine next abstractNumId
    existing_abstract = numbering.findall(qn('w:abstractNum'))
    abstract_num_id = str(len(existing_abstract))

    abstract_num = make_el('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), abstract_num_id)

    # nsid
    nsid = make_el('w:nsid')
    nsid.set(qn('w:val'), 'AB123401')
    abstract_num.append(nsid)

    # multiLevelType
    multi = make_el('w:multiLevelType')
    multi.set(qn('w:val'), 'hybridMultilevel')
    abstract_num.append(multi)

    # Define 3 levels — all using U+2022
    bullet_chars = ['\u2022', '\u2022', '\u2022']
    indent_left = [720, 1440, 2160]     # twips
    hanging = [360, 360, 360]

    for lvl_idx in range(3):
        lvl = make_el('w:lvl')
        lvl.set(qn('w:ilvl'), str(lvl_idx))

        start = make_el('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        numFmt = make_el('w:numFmt')
        numFmt.set(qn('w:val'), 'bullet')
        lvl.append(numFmt)

        lvlText = make_el('w:lvlText')
        lvlText.set(qn('w:val'), bullet_chars[lvl_idx])
        lvl.append(lvlText)

        lvlJc = make_el('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)

        pPr = make_el('w:pPr')
        ind = make_el('w:ind')
        ind.set(qn('w:left'), str(indent_left[lvl_idx]))
        ind.set(qn('w:hanging'), str(hanging[lvl_idx]))
        pPr.append(ind)
        lvl.append(pPr)

        rPr = make_el('w:rPr')
        rFonts = make_el('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Symbol')
        rFonts.set(qn('w:hAnsi'), 'Symbol')
        rFonts.set(qn('w:hint'), 'default')
        rPr.append(rFonts)
        lvl.append(rPr)

        abstract_num.append(lvl)

    numbering.append(abstract_num)

    # Create a num element that references this abstractNum
    existing_nums = numbering.findall(qn('w:num'))
    num_id_val = str(len(existing_nums) + 1)

    num = make_el('w:num')
    num.set(qn('w:numId'), num_id_val)
    abs_ref = make_el('w:abstractNumId')
    abs_ref.set(qn('w:val'), abstract_num_id)
    num.append(abs_ref)
    numbering.append(num)

    return int(num_id_val)


def set_list_para(para, num_id, ilvl):
    """Apply list numbering to a paragraph at the given indent level."""
    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')

    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_el)

    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))
    numPr.append(numId_el)

    pPr.append(numPr)


def create_initial():
    doc = Document()

    # Remove default empty paragraph
    # (Document() starts with one empty paragraph; we'll overwrite it)

    # Create the numbering definition (all levels same bullet)
    num_id = add_numbering_with_same_bullet(doc)

    # Define the content structure: (text, level 0-indexed)
    items = [
        ('User Interface Design',       0),
        ('Navigation patterns',         1),
        ('Breadcrumb navigation',       2),
        ('Hamburger menu',              2),
        ('Visual design principles',    1),
        ('Backend Architecture',        0),
        ('Microservices',               1),
        ('Service discovery',           2),
        ('Database design',             1),
    ]

    # Remove default empty paragraph(s) if any
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Add list paragraphs
    for text, lvl in items:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(12)
        set_list_para(para, num_id, lvl)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
