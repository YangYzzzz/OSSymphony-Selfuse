"""
Initial Setup: Single-column A4 document for column layout task
Task ID: writer_fs_040
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_040'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Configure A4 page with 2 cm margins (single-column default)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # --- Title ---
    title = doc.add_heading("Quarterly Market Analysis Report", level=1)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(18)

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("Prepared by the Strategic Research Division — Q1 2025")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Executive Summary ---
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(
        "The first quarter of 2025 demonstrated robust growth across all major "
        "market segments, with the technology sector leading at 8.3% quarter-over-quarter "
        "expansion. Consumer confidence indices rose to 112.4, the highest reading since "
        "mid-2023, driven by declining inflation rates and a resilient labor market. "
        "Our portfolio allocation strategy, emphasizing diversified exposure to both "
        "growth and value equities, yielded a net return of 6.7% against the benchmark "
        "of 5.2%."
    )
    doc.add_paragraph(
        "Key risk factors for Q2 include potential monetary policy tightening by the "
        "European Central Bank, ongoing supply chain disruptions in the semiconductor "
        "industry, and geopolitical tensions affecting energy commodity pricing. "
        "We recommend maintaining a cautious overweight position in healthcare and "
        "renewable energy sectors while gradually reducing exposure to highly leveraged "
        "real estate investment trusts."
    )

    # --- Market Overview ---
    doc.add_heading("Market Overview", level=2)
    doc.add_paragraph(
        "Global equity markets posted their strongest January-to-March performance "
        "in three years. The S&P 500 gained 7.1%, the MSCI Europe index advanced "
        "5.8%, and emerging market benchmarks rose 4.6% in USD terms. Fixed income "
        "markets remained under pressure, with the 10-year Treasury yield climbing "
        "from 4.15% to 4.42% over the quarter."
    )
    doc.add_paragraph(
        "The technology sector benefited from accelerating enterprise adoption of "
        "artificial intelligence infrastructure. Major cloud providers reported capital "
        "expenditure increases averaging 32% year-over-year, signaling continued "
        "investment in data center capacity. Software-as-a-service companies with "
        "AI-native product offerings saw revenue growth rates of 25-40%, significantly "
        "outpacing the broader software market."
    )

    # --- Sector Performance ---
    doc.add_heading("Sector Performance", level=2)
    doc.add_paragraph(
        "Healthcare stocks advanced 6.2% during the quarter, supported by positive "
        "clinical trial results from several mid-cap biotechnology firms. Notably, "
        "Meridian Therapeutics announced Phase III success for its novel oncology "
        "compound MRD-4517, which targets a previously undruggable protein pathway. "
        "The FDA is expected to review the application under priority designation."
    )
    doc.add_paragraph(
        "Energy sector performance was mixed, with traditional oil and gas companies "
        "declining 2.1% while renewable energy producers gained 9.4%. The divergence "
        "reflects shifting capital allocation preferences among institutional investors "
        "and new regulatory frameworks in the European Union requiring portfolio-level "
        "carbon intensity disclosure."
    )

    # --- Risk Assessment ---
    doc.add_heading("Risk Assessment", level=2)
    doc.add_paragraph(
        "Our proprietary risk model indicates elevated tail risk in credit markets, "
        "with the probability of a significant spread widening event at 18.5%, up "
        "from 12.3% at year-end 2024. Contributing factors include rising corporate "
        "leverage ratios in the commercial real estate sector, where vacancy rates "
        "in Class B office space have reached 24.7% nationally."
    )
    doc.add_paragraph(
        "Currency risk remains a consideration for internationally diversified "
        "portfolios. The US dollar index weakened 1.8% during Q1, and our models "
        "project further depreciation of 2-3% over the next two quarters. We "
        "recommend maintaining partial hedges on EUR and JPY exposures using "
        "rolling three-month forward contracts."
    )

    # --- Recommendations ---
    doc.add_heading("Recommendations", level=2)
    doc.add_paragraph(
        "Based on our analysis, we recommend the following strategic adjustments "
        "for Q2 2025:"
    )
    bullets = [
        "Increase allocation to healthcare equities by 3 percentage points, "
        "focusing on companies with late-stage pipeline catalysts.",
        "Reduce exposure to commercial REIT holdings by 5 percentage points, "
        "reallocating to residential and industrial logistics properties.",
        "Initiate a 2% position in investment-grade corporate bonds with "
        "maturities of 5-7 years to capture attractive yield spreads.",
        "Maintain the existing 8% allocation to international developed markets "
        "but shift regional emphasis from Japan to Northern Europe.",
        "Continue holding a 4% cash reserve as a tactical buffer against "
        "potential market dislocations in the second half of the year."
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
