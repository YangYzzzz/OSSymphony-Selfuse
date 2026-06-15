"""
Reward Script: Add total page count field to footer in 'Page X of Y' format
Task ID: writer_legal_089
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35) — NUMPAGES field code exists in footer
  Component 2 (0.35) — 'Page ' prefix text before PAGE field
  Component 3 (0.30) — ' of ' separator text between PAGE and NUMPAGES fields
"""

import os
from docx import Document
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_089'
WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': WML_NS}


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def get_footer_xml_elements(doc):
    """Extract all runs from the first footer paragraph of the first section."""
    section = doc.sections[0]
    footer = section.footer
    if not footer.paragraphs:
        return []
    para = footer.paragraphs[0]
    return list(para._element)


def parse_footer_structure(elements):
    """
    Parse footer paragraph XML elements into a sequence of tokens.
    Each token is either:
      ('text', <string>)       — a w:t text run
      ('field', <field_name>)  — a field code like 'PAGE' or 'NUMPAGES'
    """
    tokens = []
    in_field = False
    for elem in elements:
        tag = etree.QName(elem.tag).localname if isinstance(elem.tag, str) else ''
        if tag == 'pPr':
            continue  # skip paragraph properties
        if tag == 'r':
            # Process children of this run
            for child in elem:
                child_tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ''
                if child_tag == 't':
                    text_val = child.text or ''
                    if text_val:
                        tokens.append(('text', text_val))
                elif child_tag == 'fldChar':
                    fld_type = child.get(f'{{{WML_NS}}}fldCharType')
                    if fld_type == 'begin':
                        in_field = True
                    elif fld_type == 'end':
                        in_field = False
                elif child_tag == 'instrText' and in_field:
                    field_name = (child.text or '').strip()
                    tokens.append(('field', field_name))
    return tokens


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Parse footer structure
    try:
        elements = get_footer_xml_elements(doc)
        tokens = parse_footer_structure(elements)
        print(f"DEBUG: Footer tokens: {tokens}")
    except Exception as e:
        print(f"ERROR: Could not parse footer XML: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: NUMPAGES field code exists in footer (0.35 points)
    # Initial has only PAGE field; golden adds NUMPAGES field.
    try:
        field_names = [t[1] for t in tokens if t[0] == 'field']
        if 'NUMPAGES' in field_names:
            print(f"PASS: Component 1 — NUMPAGES field found in footer (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 — NUMPAGES field not found. Fields present: {field_names}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 'Page ' prefix text before the PAGE field (0.35 points)
    # Initial has no text runs, just the bare PAGE field. Golden has 'Page ' before PAGE.
    try:
        # Find the index of the PAGE field token
        page_idx = None
        for idx, t in enumerate(tokens):
            if t[0] == 'field' and t[1] == 'PAGE':
                page_idx = idx
                break

        if page_idx is not None and page_idx > 0:
            # Check if there's a text token before PAGE containing 'Page'
            preceding_texts = [t[1] for t in tokens[:page_idx] if t[0] == 'text']
            combined_preceding = ''.join(preceding_texts).strip().lower()
            if 'page' in combined_preceding:
                print(f"PASS: Component 2 — 'Page' prefix found before PAGE field (0.35 pts)")
                total_score += 0.35
            else:
                print(f"FAIL: Component 2 — No 'Page' prefix before PAGE field. Preceding text: {repr(combined_preceding)}")
        else:
            print(f"FAIL: Component 2 — PAGE field not found or no text before it. page_idx={page_idx}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: ' of ' separator text between PAGE and NUMPAGES fields (0.30 points)
    # Initial has no NUMPAGES field, so no ' of ' separator. Golden has ' of ' between them.
    try:
        page_idx = None
        numpages_idx = None
        for idx, t in enumerate(tokens):
            if t[0] == 'field' and t[1] == 'PAGE' and page_idx is None:
                page_idx = idx
            if t[0] == 'field' and t[1] == 'NUMPAGES' and numpages_idx is None:
                numpages_idx = idx

        if page_idx is not None and numpages_idx is not None and numpages_idx > page_idx:
            # Check for 'of' text between the two fields
            between_texts = [t[1] for t in tokens[page_idx+1:numpages_idx] if t[0] == 'text']
            combined_between = ''.join(between_texts).strip().lower()
            if 'of' in combined_between:
                print(f"PASS: Component 3 — 'of' separator found between PAGE and NUMPAGES (0.30 pts)")
                total_score += 0.30
            else:
                print(f"FAIL: Component 3 — No 'of' separator between PAGE and NUMPAGES. Between text: {repr(combined_between)}")
        else:
            print(f"FAIL: Component 3 — Cannot find both PAGE and NUMPAGES fields in order. PAGE={page_idx}, NUMPAGES={numpages_idx}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
