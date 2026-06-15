"""
Initial Setup: Create a legal document with footer showing only page number.
Task ID: writer_legal_089
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_089'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_field(paragraph):
    """Add a PAGE field code to the paragraph (current page number only)."""
    # Begin field
    r1 = paragraph.add_run()
    fldChar1 = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fldChar1)

    # Instruction
    r2 = paragraph.add_run()
    instrText = r2._element.makeelement(qn('w:instrText'), {})
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    r2._element.append(instrText)

    # End field
    r3 = paragraph.add_run()
    fldChar3 = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fldChar3)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Footer: page number only ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number_field(fp)

    # --- Document Title ---
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Parties
    doc.add_heading('1. PARTIES', level=1)
    p = doc.add_paragraph()
    p.add_run('This Professional Services Agreement ("Agreement") is entered into as of March 15, 2025, by and between:')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Meridian Technology Solutions, Inc.').bold = True
    p.add_run(', a Delaware corporation with principal offices at 2400 Innovation Drive, Suite 800, San Jose, California 95134 ("Company"); and')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Clearwater Consulting Group, LLC').bold = True
    p.add_run(', a California limited liability company with principal offices at 1750 Montgomery Street, Suite 450, San Francisco, California 94111 ("Consultant").')

    # Scope of Services
    doc.add_heading('2. SCOPE OF SERVICES', level=1)
    doc.add_paragraph(
        'The Consultant agrees to provide the following professional services to the Company '
        'in connection with the implementation of the Enterprise Resource Planning (ERP) system '
        'migration project (the "Project"):'
    )
    services = [
        'Comprehensive assessment of existing IT infrastructure and legacy systems, including hardware inventory, software licensing audit, and network topology mapping.',
        'Design and architecture of the target ERP environment, incorporating cloud-native deployment strategies, microservices patterns, and API-first integration approaches.',
        'Data migration planning and execution, including extraction of historical records from the legacy Oracle Database 11g system, transformation logic development, and loading into the new SAP S/4HANA Cloud environment.',
        'Custom development of integration middleware connecting the ERP system with existing CRM (Salesforce), HCM (Workday), and financial reporting (Adaptive Insights) platforms.',
        'End-user training program development and delivery for approximately 350 employees across 12 departments, including role-based training modules and competency assessments.',
        'Post-implementation support for a period of ninety (90) calendar days, including bug resolution, performance tuning, and knowledge transfer to the Company\'s internal IT team.',
    ]
    for svc in services:
        doc.add_paragraph(svc, style='List Bullet')

    # Term and Termination
    doc.add_heading('3. TERM AND TERMINATION', level=1)
    doc.add_paragraph(
        'This Agreement shall commence on April 1, 2025, and shall continue for a period of '
        'eighteen (18) months unless earlier terminated in accordance with the provisions of this Section.'
    )
    doc.add_paragraph(
        '3.1 Termination for Convenience. Either party may terminate this Agreement upon sixty (60) days\' '
        'prior written notice to the other party. In the event of termination for convenience by the Company, '
        'the Company shall pay the Consultant for all Services performed and expenses incurred through the effective date of termination.'
    )
    doc.add_paragraph(
        '3.2 Termination for Cause. Either party may terminate this Agreement immediately upon written notice '
        'if the other party: (a) materially breaches any provision of this Agreement and fails to cure such breach '
        'within thirty (30) days after receipt of written notice; (b) becomes insolvent or files a petition for '
        'bankruptcy; or (c) ceases to conduct business in the normal course.'
    )

    # Compensation
    doc.add_heading('4. COMPENSATION AND PAYMENT TERMS', level=1)
    doc.add_paragraph(
        '4.1 Professional Fees. The Company shall pay the Consultant the following fees for Services rendered:'
    )
    # Fee table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Service Category', 'Rate', 'Estimated Hours']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    data = [
        ['Infrastructure Assessment', '$275/hour', '160'],
        ['Architecture & Design', '$325/hour', '240'],
        ['Data Migration', '$300/hour', '320'],
        ['Integration Development', '$350/hour', '480'],
        ['Training & Documentation', '$225/hour', '200'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_paragraph(
        '4.2 Payment Terms. The Consultant shall submit monthly invoices to the Company detailing '
        'the Services performed during the preceding calendar month. Each invoice shall include a '
        'breakdown of hours worked by service category, a description of deliverables completed, '
        'and any reimbursable expenses incurred. The Company shall pay each undisputed invoice '
        'within thirty (30) days of receipt.'
    )
    doc.add_paragraph(
        '4.3 Expenses. The Company shall reimburse the Consultant for reasonable and pre-approved '
        'travel, lodging, and incidental expenses incurred in connection with the performance of '
        'Services. All expenses exceeding $500.00 require prior written approval from the Company\'s designated project manager.'
    )

    # Confidentiality
    doc.add_heading('5. CONFIDENTIALITY', level=1)
    doc.add_paragraph(
        '5.1 Definition. "Confidential Information" means all non-public information disclosed by either '
        'party to the other, whether orally, in writing, or by inspection of tangible objects, including '
        'without limitation: trade secrets, business plans, financial data, customer lists, technical '
        'specifications, source code, algorithms, product roadmaps, marketing strategies, and employee information.'
    )
    doc.add_paragraph(
        '5.2 Obligations. Each party agrees to: (a) hold the other party\'s Confidential Information in '
        'strict confidence; (b) not disclose such information to any third party without prior written consent; '
        '(c) use such information solely for the purposes of performing obligations under this Agreement; and '
        '(d) protect such information using the same degree of care it uses to protect its own confidential '
        'information, but in no event less than reasonable care.'
    )
    doc.add_paragraph(
        '5.3 Survival. The confidentiality obligations set forth in this Section shall survive the '
        'expiration or termination of this Agreement for a period of three (3) years.'
    )

    # Intellectual Property
    doc.add_heading('6. INTELLECTUAL PROPERTY RIGHTS', level=1)
    doc.add_paragraph(
        '6.1 Work Product. All work product, deliverables, inventions, discoveries, and improvements '
        'created by the Consultant in the course of performing Services under this Agreement ("Work Product") '
        'shall be the sole and exclusive property of the Company. The Consultant hereby assigns to the Company '
        'all right, title, and interest in and to such Work Product, including all intellectual property rights therein.'
    )
    doc.add_paragraph(
        '6.2 Pre-Existing Materials. The Consultant retains all rights in any tools, methodologies, '
        'frameworks, and other materials owned by or licensed to the Consultant prior to the commencement '
        'of this Agreement ("Pre-Existing Materials"). To the extent any Pre-Existing Materials are incorporated '
        'into any Work Product, the Consultant grants the Company a perpetual, irrevocable, worldwide, '
        'royalty-free license to use, modify, and distribute such Pre-Existing Materials as part of the Work Product.'
    )

    # Limitation of Liability
    doc.add_heading('7. LIMITATION OF LIABILITY', level=1)
    doc.add_paragraph(
        '7.1 IN NO EVENT SHALL EITHER PARTY BE LIABLE TO THE OTHER PARTY FOR ANY INDIRECT, INCIDENTAL, '
        'SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES ARISING OUT OF OR RELATED TO THIS AGREEMENT, REGARDLESS '
        'OF WHETHER SUCH DAMAGES ARE BASED ON CONTRACT, TORT, STRICT LIABILITY, OR ANY OTHER THEORY, EVEN IF '
        'THE PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.'
    )
    doc.add_paragraph(
        '7.2 THE TOTAL CUMULATIVE LIABILITY OF EITHER PARTY UNDER THIS AGREEMENT SHALL NOT EXCEED THE TOTAL '
        'FEES PAID OR PAYABLE BY THE COMPANY TO THE CONSULTANT DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY '
        'PRECEDING THE EVENT GIVING RISE TO THE CLAIM.'
    )

    # Governing Law
    doc.add_heading('8. GOVERNING LAW AND DISPUTE RESOLUTION', level=1)
    doc.add_paragraph(
        '8.1 Governing Law. This Agreement shall be governed by and construed in accordance with the laws '
        'of the State of California, without regard to its conflicts of laws principles.'
    )
    doc.add_paragraph(
        '8.2 Dispute Resolution. Any dispute arising out of or relating to this Agreement shall first be '
        'submitted to good faith mediation. If mediation is unsuccessful, the dispute shall be resolved by '
        'binding arbitration administered by the American Arbitration Association in accordance with its '
        'Commercial Arbitration Rules. The arbitration shall be conducted in San Francisco, California, '
        'before a single arbitrator mutually agreed upon by the parties.'
    )

    # Miscellaneous
    doc.add_heading('9. MISCELLANEOUS PROVISIONS', level=1)
    doc.add_paragraph(
        '9.1 Entire Agreement. This Agreement constitutes the entire agreement between the parties with '
        'respect to the subject matter hereof and supersedes all prior and contemporaneous agreements, '
        'understandings, negotiations, and discussions, whether oral or written.'
    )
    doc.add_paragraph(
        '9.2 Amendment. No modification or amendment of this Agreement shall be valid unless made in '
        'writing and signed by both parties.'
    )
    doc.add_paragraph(
        '9.3 Waiver. The failure of either party to enforce any provision of this Agreement shall not '
        'constitute a waiver of such provision or the right to enforce it at a later time.'
    )
    doc.add_paragraph(
        '9.4 Severability. If any provision of this Agreement is held to be invalid, illegal, or '
        'unenforceable, the remaining provisions shall continue in full force and effect.'
    )
    doc.add_paragraph(
        '9.5 Assignment. Neither party may assign this Agreement or any rights or obligations hereunder '
        'without the prior written consent of the other party, except that the Company may assign this '
        'Agreement to an affiliate or in connection with a merger, acquisition, or sale of all or '
        'substantially all of its assets.'
    )

    # Signature block
    doc.add_paragraph()
    doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.')
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('MERIDIAN TECHNOLOGY SOLUTIONS, INC.').bold = True
    doc.add_paragraph()
    doc.add_paragraph('By: ___________________________')
    doc.add_paragraph('Name: Victoria R. Hartwell')
    doc.add_paragraph('Title: Chief Technology Officer')
    doc.add_paragraph('Date: March 15, 2025')
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('CLEARWATER CONSULTING GROUP, LLC').bold = True
    doc.add_paragraph()
    doc.add_paragraph('By: ___________________________')
    doc.add_paragraph('Name: David M. Nakamura')
    doc.add_paragraph('Title: Managing Partner')
    doc.add_paragraph('Date: March 15, 2025')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
