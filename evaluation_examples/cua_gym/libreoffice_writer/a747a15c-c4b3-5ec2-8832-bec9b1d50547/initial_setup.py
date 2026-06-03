"""
Initial Setup: Anti-Harassment Policy document with all Default Paragraph Style
Task ID: writer_hr_040
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_040'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Title (Normal style, NOT heading) ---
    p = doc.add_paragraph('Anti-Harassment Policy')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(16)

    # --- Introductory paragraph ---
    doc.add_paragraph(
        'This policy outlines the commitment of Meridian Global Solutions to maintaining '
        'a workplace free from harassment. All employees, contractors, and visitors are '
        'expected to comply with this policy. Violations will be addressed promptly and '
        'may result in disciplinary action up to and including termination of employment.'
    )

    doc.add_paragraph(
        'Meridian Global Solutions recognizes that harassment undermines workplace dignity '
        'and productivity. This policy applies to all work-related settings including offices, '
        'client sites, work-related social events, and electronic communications.'
    )

    # ========== Section 1: Definition ==========
    p = doc.add_paragraph('Definition')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(13)

    doc.add_paragraph(
        'Harassment is any unwelcome conduct based on a protected characteristic that creates '
        'an intimidating, hostile, or offensive work environment, or that unreasonably interferes '
        'with an individual\'s work performance.'
    )

    # Sub-subsection 1.1
    p = doc.add_paragraph('Verbal Harassment')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph(
        'Verbal harassment includes derogatory comments, slurs, jokes, or epithets directed at '
        'an individual based on race, gender, religion, national origin, age, disability, or '
        'sexual orientation. This also covers repeated unwelcome remarks about a person\'s '
        'appearance, lifestyle, or personal beliefs that cause discomfort or distress.'
    )

    # Sub-subsection 1.2
    p = doc.add_paragraph('Physical Harassment')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph(
        'Physical harassment encompasses unwanted physical contact, blocking movement, or '
        'intimidating physical gestures. This includes but is not limited to pushing, grabbing, '
        'cornering, or any form of assault. Even seemingly minor physical contact can constitute '
        'harassment if it is unwelcome and related to a protected characteristic.'
    )

    # Sub-subsection 1.3
    p = doc.add_paragraph('Digital Harassment')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph(
        'Digital harassment includes sending offensive emails, instant messages, or social media '
        'posts. It also covers sharing inappropriate images, cyberstalking, or creating a hostile '
        'online environment through work communication platforms such as Slack, Microsoft Teams, '
        'or company email systems.'
    )

    # ========== Section 2: Reporting Procedures ==========
    p = doc.add_paragraph('Reporting Procedures')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(13)

    doc.add_paragraph(
        'Employees who experience or witness harassment are encouraged to report the incident '
        'as soon as possible. The company provides multiple reporting channels to ensure '
        'accessibility and confidentiality.'
    )

    # Sub-subsection 2.1
    p = doc.add_paragraph('Internal Reporting Channels')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph(
        'Reports can be submitted to the employee\'s direct supervisor, the Human Resources '
        'department, or through the anonymous ethics hotline at 1-800-555-0142. The HR department '
        'maintains a dedicated intake form available on the company intranet under the Employee '
        'Relations section. All reports are logged in a secure case management system.'
    )

    # Sub-subsection 2.2
    p = doc.add_paragraph('External Reporting Options')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph(
        'Employees may also file complaints with the Equal Employment Opportunity Commission (EEOC) '
        'or the relevant state civil rights agency. The company will cooperate fully with any '
        'external investigation. Contact information for the EEOC regional office is posted in '
        'all break rooms and on the HR intranet page.'
    )

    # Sub-subsection 2.3
    p = doc.add_paragraph('Confidentiality Protections')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph(
        'All reports of harassment will be treated with the highest level of confidentiality '
        'consistent with a thorough investigation. Information will be shared only on a '
        'need-to-know basis. Retaliation against any individual who reports harassment or '
        'participates in an investigation is strictly prohibited and will result in separate '
        'disciplinary action.'
    )

    # ========== Section 3: Investigation Process ==========
    p = doc.add_paragraph('Investigation Process')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(13)

    doc.add_paragraph(
        'Upon receipt of a harassment complaint, the company will initiate a prompt, thorough, '
        'and impartial investigation. The investigation will be conducted by trained HR personnel '
        'or an external investigator as appropriate.'
    )

    # Sub-subsection 3.1
    p = doc.add_paragraph('Initial Assessment')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph(
        'Within 48 hours of receiving a complaint, the HR department will conduct an initial '
        'assessment to determine the scope and severity of the allegations. This may include '
        'a preliminary interview with the complainant and a review of any immediately available '
        'evidence such as emails, chat logs, or security footage.'
    )

    # Sub-subsection 3.2
    p = doc.add_paragraph('Formal Investigation Steps')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph(
        'The formal investigation includes interviewing all relevant parties, collecting documentary '
        'evidence, reviewing personnel records, and consulting with legal counsel as needed. '
        'Interviews are conducted individually and in private. Written statements are obtained '
        'from all witnesses. The investigation will typically be completed within 30 business days.'
    )

    # ========== Section 4: Consequences ==========
    p = doc.add_paragraph('Consequences')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(13)

    doc.add_paragraph(
        'Employees found to have engaged in harassment will face disciplinary action proportionate '
        'to the severity of the offense. The company reserves the right to impose any level of '
        'discipline it deems appropriate.'
    )

    # Sub-subsection 4.1
    p = doc.add_paragraph('Disciplinary Actions')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph(
        'Disciplinary measures may include verbal warning, written warning, mandatory harassment '
        'prevention training, suspension without pay, demotion, transfer, or termination. The '
        'specific action will depend on the nature and severity of the harassment, the offender\'s '
        'history, and the impact on the affected individual.'
    )

    # Sub-subsection 4.2
    p = doc.add_paragraph('Repeat Offenses')
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph(
        'Second or subsequent substantiated complaints against the same individual will result '
        'in escalated disciplinary action. A pattern of harassment behavior, even if each individual '
        'incident might warrant lesser discipline, will be treated as a serious offense potentially '
        'resulting in immediate termination. All disciplinary records are maintained in the employee\'s '
        'permanent personnel file.'
    )

    # Final note
    doc.add_paragraph(
        'This policy is reviewed annually by the Legal and Human Resources departments. '
        'Last updated: March 2025. Next scheduled review: March 2026. Questions regarding '
        'this policy should be directed to the HR Policy team at hr-policy@meridianglobal.com.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
