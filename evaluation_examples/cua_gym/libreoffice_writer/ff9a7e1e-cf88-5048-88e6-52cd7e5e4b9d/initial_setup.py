"""
Initial Setup: Configure AutoCorrect to replace ':)' with smiley emoji
Task ID: writer_frd_057
Domain: libreoffice_writer

Initial state: LibreOffice Writer open with a casual document.
No custom AutoCorrect entry for ':)' exists.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_057'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("Weekend Plans with Friends", level=1)

    # Intro paragraph
    p1 = doc.add_paragraph()
    run1 = p1.add_run("Hey everyone! Just wanted to jot down some ideas for our "
                       "get-together this Saturday. Let me know what you think!")
    run1.font.size = Pt(12)
    run1.font.name = "Calibri"

    # Activities section
    doc.add_heading("Activity Options", level=2)

    p2 = doc.add_paragraph()
    run2 = p2.add_run("Here are a few things we could do:")
    run2.font.size = Pt(12)
    run2.font.name = "Calibri"

    activities = [
        "Morning hike at Cedar Ridge Trail - the weather looks perfect for it",
        "Brunch at The Maple Leaf Cafe - they have that new avocado toast everyone's been talking about",
        "Afternoon board games at my place - I just got Settlers of Catan and Ticket to Ride",
        "Evening movie marathon - thinking comedy or sci-fi, open to suggestions",
        "BBQ in the backyard if the weather holds up",
    ]
    for activity in activities:
        bp = doc.add_paragraph(activity, style="List Bullet")
        for run in bp.runs:
            run.font.size = Pt(12)
            run.font.name = "Calibri"

    # Who's coming section
    doc.add_heading("Who's Coming?", level=2)

    attendees_text = (
        "So far confirmed: Alex, Jamie, Morgan, and Taylor. "
        "Still waiting to hear back from Sam and Jordan. "
        "If anyone has dietary restrictions for the BBQ, please let me know ASAP!"
    )
    p3 = doc.add_paragraph()
    run3 = p3.add_run(attendees_text)
    run3.font.size = Pt(12)
    run3.font.name = "Calibri"

    # Note about bring items
    doc.add_heading("What to Bring", level=2)

    items = [
        "Sunscreen and hats for the hike",
        "Your favorite board game if you have one",
        "A dish or drinks to share for the BBQ",
        "Blankets for the movie marathon - my living room gets chilly",
    ]
    for item in items:
        bp = doc.add_paragraph(item, style="List Bullet")
        for run in bp.runs:
            run.font.size = Pt(12)
            run.font.name = "Calibri"

    # Closing
    p_close = doc.add_paragraph()
    run_close = p_close.add_run(
        "Can't wait to see you all! It's going to be a great day. "
        "Text me if you have any questions or want to suggest something else."
    )
    run_close.font.size = Pt(12)
    run_close.font.name = "Calibri"

    p_sign = doc.add_paragraph()
    run_sign = p_sign.add_run("- Riley")
    run_sign.font.size = Pt(12)
    run_sign.font.name = "Calibri"
    run_sign.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
