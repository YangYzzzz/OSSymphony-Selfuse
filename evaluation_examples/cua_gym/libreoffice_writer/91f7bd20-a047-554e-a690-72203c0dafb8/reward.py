"""
Reward Script: Thesis title page creation
Task ID: writer_acad_030
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Title text present and centered
  Component 2 (0.20): Title is 16pt bold
  Component 3 (0.35): Author, department, university, date present and centered
  Component 4 (0.10): Author/dept/uni/date are 12pt
  Component 5 (0.10): No header/footer content (clean title page)
"""

import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_030'

# Persistence hook: save any unsaved GUI state before verification
def persist_app_state():
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui, time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def get_non_empty_paragraphs(doc):
    """Return list of paragraphs that have non-empty text."""
    return [p for p in doc.paragraphs if p.text.strip()]


def verify_task(file_path):
    """
    Verify thesis title page creation with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    non_empty = get_non_empty_paragraphs(doc)

    # We expect at least 5 non-empty paragraphs: title, author, department, university, date
    # The task says "each on a separate line"

    # Component 1: Title text present and centered (0.25 points)
    # The title must be "Machine Learning Approaches to Climate Prediction" and centered
    try:
        title_found = False
        for para in doc.paragraphs:
            if 'Machine Learning Approaches to Climate Prediction' in para.text:
                if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    print(f"PASS: Component 1 — Title found and centered (0.25 pts)")
                    total_score += 0.25
                    title_found = True
                else:
                    print(f"FAIL: Component 1 — Title found but not centered (alignment={para.paragraph_format.alignment})")
                    title_found = True
                break
        if not title_found:
            print("FAIL: Component 1 — Title text 'Machine Learning Approaches to Climate Prediction' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Title is 16pt bold (0.20 points)
    try:
        title_format_ok = False
        for para in doc.paragraphs:
            if 'Machine Learning Approaches to Climate Prediction' in para.text:
                for run in para.runs:
                    if 'Machine Learning' in run.text:
                        is_bold = run.font.bold is True
                        size_ok = (run.font.size is not None and
                                   abs(run.font.size.pt - 16.0) < 0.5)
                        if is_bold and size_ok:
                            print(f"PASS: Component 2 — Title is 16pt bold (0.20 pts)")
                            total_score += 0.20
                            title_format_ok = True
                        else:
                            print(f"FAIL: Component 2 — Title bold={run.font.bold}, size={run.font.size.pt if run.font.size else None}")
                        break
                break
        if not title_format_ok and total_score < 0.25:
            # Already printed fail or title not found
            pass
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Author, department, university, date present and centered (0.35 points)
    # Each item is worth 0.35/4 = 0.0875 points
    expected_items = [
        ('John Smith', 'Author'),
        ('Department of Computer Science', 'Department'),
        ('Stanford University', 'University'),
        ('June 2025', 'Date'),
    ]
    try:
        comp3_score = 0.0
        per_item = 0.35 / 4.0
        for expected_text, label in expected_items:
            found = False
            for para in doc.paragraphs:
                if expected_text in para.text:
                    if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                        comp3_score += per_item
                        print(f"  PASS: {label} '{expected_text}' found and centered")
                        found = True
                    else:
                        print(f"  FAIL: {label} '{expected_text}' found but not centered")
                        found = True
                    break
            if not found:
                print(f"  FAIL: {label} '{expected_text}' not found")
        if comp3_score > 0:
            print(f"PASS: Component 3 — {comp3_score:.4f}/{0.35} pts for info items")
            total_score += comp3_score
        else:
            print(f"FAIL: Component 3 — No info items found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Author/dept/uni/date text is ~12pt (0.10 points)
    try:
        comp4_items_ok = 0
        for expected_text, label in expected_items:
            for para in doc.paragraphs:
                if expected_text in para.text:
                    for run in para.runs:
                        if expected_text in run.text or run.text.strip() in expected_text:
                            if run.font.size is not None and abs(run.font.size.pt - 12.0) < 0.5:
                                comp4_items_ok += 1
                            break
                    break
        if comp4_items_ok >= 3:
            print(f"PASS: Component 4 — {comp4_items_ok}/4 info items are 12pt (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — Only {comp4_items_ok}/4 info items at 12pt")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Title page has content AND no visible header/footer/page number (0.10 points)
    # This is a compound check: title page must have the thesis content (task-introduced)
    # AND must not have headers/footers/page numbers. Both conditions required.
    try:
        has_title_content = any('Machine Learning Approaches to Climate Prediction' in p.text for p in doc.paragraphs)

        section = doc.sections[0]
        header_empty = all(p.text.strip() == '' for p in section.header.paragraphs)
        footer_empty = all(p.text.strip() == '' for p in section.footer.paragraphs)

        # Also check XML for page number field codes in footer
        from docx.oxml.ns import qn
        footer_has_page_num = False
        for p in section.footer.paragraphs:
            for run in p.runs:
                if run.element.findall('.//' + qn('w:fldChar')):
                    footer_has_page_num = True
                    break
                instr_elems = run.element.findall('.//' + qn('w:instrText'))
                for instr in instr_elems:
                    if 'PAGE' in (instr.text or ''):
                        footer_has_page_num = True
                        break

        if has_title_content and header_empty and footer_empty and not footer_has_page_num:
            print(f"PASS: Component 5 — Title page content present with no header/footer/page number (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — content={has_title_content}, header_empty={header_empty}, footer_empty={footer_empty}, page_num={footer_has_page_num}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state()
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
