"""
Initial Setup: Format academic paper with structured headings and TOC
Task ID: writer_pd_006
Domain: libreoffice_writer

Creates a 45-page academic paper with chapter/section/subsection titles
styled as bold/larger font using Normal (Default Paragraph Style).
No Heading styles or TOC applied -- those are the agent's task.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# Body text paragraphs to fill pages (each ~3-5 sentences of academic prose)
BODY_TEXTS = [
    "The rapid advancement of machine learning techniques has fundamentally transformed how researchers approach complex computational problems. Deep neural networks, in particular, have demonstrated remarkable capabilities across diverse application domains, ranging from natural language processing to computer vision. These developments have been accompanied by significant theoretical insights into the nature of representation learning and optimization landscapes.",
    "Statistical methods form the backbone of empirical research in the computational sciences. The application of rigorous hypothesis testing, combined with Bayesian inference frameworks, enables researchers to draw meaningful conclusions from experimental data. Recent advances in causal inference have further strengthened the methodological toolkit available to practitioners.",
    "Distributed computing architectures have evolved to support the increasing computational demands of modern research. Cloud-based infrastructure provides scalable resources for training large-scale models, while edge computing paradigms enable real-time inference in resource-constrained environments. The interplay between these computing paradigms continues to shape the landscape of computational research.",
    "Data collection and curation represent critical challenges in empirical research. The quality of training datasets directly impacts model performance, necessitating careful attention to sampling strategies, annotation protocols, and bias mitigation techniques. Recent work has highlighted the importance of dataset documentation and provenance tracking.",
    "Optimization algorithms play a central role in training modern machine learning models. Stochastic gradient descent and its variants remain the workhorses of deep learning, while second-order methods and natural gradient approaches offer theoretical advantages in certain settings. The convergence properties of these algorithms continue to be an active area of research.",
    "Transfer learning has emerged as a powerful paradigm for leveraging knowledge across related tasks and domains. Pre-trained models, particularly large language models, have demonstrated remarkable few-shot learning capabilities. The theoretical foundations of transfer learning, including domain adaptation theory, provide insights into when and why transfer succeeds.",
    "The ethical implications of computational research have received increasing attention from the academic community. Fairness, accountability, and transparency in algorithmic decision-making are now recognized as fundamental requirements rather than optional considerations. Frameworks for responsible AI development continue to evolve in response to emerging challenges.",
    "Experimental design in computational research requires careful consideration of evaluation metrics, baseline comparisons, and statistical significance. Reproducibility has emerged as a central concern, with the research community developing standards for code sharing, data availability, and experimental documentation. These efforts aim to strengthen the scientific foundations of the field.",
    "Graph-based models have gained prominence for reasoning about relational and structured data. Graph neural networks extend deep learning techniques to non-Euclidean domains, enabling applications in molecular chemistry, social network analysis, and knowledge graph reasoning. The expressiveness and limitations of these models are subjects of ongoing theoretical investigation.",
    "Reinforcement learning algorithms enable agents to learn optimal strategies through interaction with complex environments. Model-free approaches such as policy gradient methods have achieved notable successes in game playing and robotic control. Model-based reinforcement learning offers improved sample efficiency by leveraging learned dynamics models.",
    "Natural language understanding requires capturing the nuanced relationships between words, phrases, and discourse structures. Transformer architectures have revolutionized this field through self-attention mechanisms that model long-range dependencies efficiently. Multilingual models extend these capabilities across language boundaries, raising questions about universal linguistic representations.",
    "Computer vision research has progressed from handcrafted feature extraction to end-to-end learned representations. Convolutional neural networks established the foundation for modern visual recognition systems, while vision transformers have demonstrated competitive or superior performance on many benchmarks. Self-supervised learning from unlabeled visual data represents a promising direction for reducing annotation costs.",
    "The intersection of computational methods and biological sciences has opened new frontiers in drug discovery, genomics, and protein structure prediction. AlphaFold and related approaches have demonstrated that deep learning can solve longstanding problems in structural biology. These successes motivate further investigation into the application of computational methods across the life sciences.",
    "Privacy-preserving computation techniques address the growing need to analyze sensitive data while maintaining individual privacy guarantees. Differential privacy provides formal mathematical frameworks for quantifying privacy loss, while federated learning enables collaborative model training without centralizing raw data. The tension between utility and privacy remains an active area of research.",
    "Robustness and reliability of machine learning systems are essential for real-world deployment. Adversarial examples reveal vulnerabilities in neural networks, motivating research into certified defenses and robust optimization. Distribution shift, where test data differs from training data, poses additional challenges that require careful handling in production systems.",
]


def add_body_paragraphs(doc, count=3, start_idx=0):
    """Add realistic body paragraphs to fill content."""
    for i in range(count):
        idx = (start_idx + i) % len(BODY_TEXTS)
        p = doc.add_paragraph(BODY_TEXTS[idx])
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'


def add_chapter_title(doc, title):
    """Add a chapter title styled as bold, 18pt but using Normal style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.page_break_before = True
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'
    return p


def add_section_title(doc, title):
    """Add a section title styled as bold, 15pt but using Normal style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(15)
    run.font.name = 'Times New Roman'
    return p


def add_subsection_title(doc, title):
    """Add a subsection title styled as bold italic, 13pt but using Normal style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.bold = True
    run.italic = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    return p


def create_initial():
    doc = Document()

    # Set default page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ===================== PAGE 1: TITLE PAGE =====================
    # Add some spacing before the title
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_p.add_run('Advances in Computational Methods for\nMulti-Domain Scientific Discovery')
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.name = 'Times New Roman'

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_p.paragraph_format.space_before = Pt(24)
    sub_run = subtitle_p.add_run('A Doctoral Thesis')
    sub_run.font.size = Pt(16)
    sub_run.font.name = 'Times New Roman'
    sub_run.italic = True

    author_p = doc.add_paragraph()
    author_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_p.paragraph_format.space_before = Pt(36)
    auth_run = author_p.add_run('Presented by\nDr. Elena Vasquez-Richardson')
    auth_run.font.size = Pt(14)
    auth_run.font.name = 'Times New Roman'

    dept_p = doc.add_paragraph()
    dept_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept_p.paragraph_format.space_before = Pt(24)
    dept_run = dept_p.add_run('Department of Computer Science and Engineering\nStanford University\n2025')
    dept_run.font.size = Pt(12)
    dept_run.font.name = 'Times New Roman'

    # ===================== PAGE 2: BLANK (for TOC) =====================
    doc.add_page_break()
    toc_placeholder = doc.add_paragraph()
    toc_placeholder.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    toc_placeholder.paragraph_format.space_before = Pt(200)
    toc_run = toc_placeholder.add_run('[Table of Contents will be inserted here]')
    toc_run.font.size = Pt(12)
    toc_run.font.name = 'Times New Roman'
    toc_run.italic = True
    toc_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ===================== CHAPTER 1: Introduction =====================
    # Chapter 1 (Heading 1 candidate)
    add_chapter_title(doc, 'Chapter 1: Introduction')

    # Section 1.1 (Heading 2 candidate)
    add_section_title(doc, '1.1 Background and Motivation')
    add_body_paragraphs(doc, 4, start_idx=0)

    # Section 1.2
    add_section_title(doc, '1.2 Research Objectives')
    add_body_paragraphs(doc, 3, start_idx=4)

    # Subsection 1.2.1 (Heading 3 candidate)
    add_subsection_title(doc, '1.2.1 Primary Research Questions')
    add_body_paragraphs(doc, 3, start_idx=7)

    # Subsection 1.2.2
    add_subsection_title(doc, '1.2.2 Scope and Limitations')
    add_body_paragraphs(doc, 3, start_idx=10)

    # Section 1.3
    add_section_title(doc, '1.3 Thesis Organization')
    add_body_paragraphs(doc, 3, start_idx=13)

    # ===================== CHAPTER 2: Literature Review =====================
    add_chapter_title(doc, 'Chapter 2: Literature Review')

    add_section_title(doc, '2.1 Foundations of Machine Learning')
    add_body_paragraphs(doc, 4, start_idx=1)

    add_subsection_title(doc, '2.1.1 Supervised Learning Paradigms')
    add_body_paragraphs(doc, 3, start_idx=5)

    add_subsection_title(doc, '2.1.2 Unsupervised and Self-Supervised Methods')
    add_body_paragraphs(doc, 3, start_idx=9)

    add_section_title(doc, '2.2 Deep Learning Architectures')
    add_body_paragraphs(doc, 4, start_idx=3)

    # Additional body content for chapter 2
    add_body_paragraphs(doc, 3, start_idx=7)

    # ===================== CHAPTER 3: Methodology =====================
    add_chapter_title(doc, 'Chapter 3: Research Methodology')

    add_section_title(doc, '3.1 Experimental Framework')
    add_body_paragraphs(doc, 4, start_idx=2)

    add_subsection_title(doc, '3.1.1 Data Collection Protocol')
    add_body_paragraphs(doc, 3, start_idx=6)

    add_section_title(doc, '3.2 Model Architecture Design')
    add_body_paragraphs(doc, 4, start_idx=10)

    add_section_title(doc, '3.3 Evaluation Metrics')
    add_body_paragraphs(doc, 3, start_idx=0)

    add_subsection_title(doc, '3.3.1 Quantitative Assessment Criteria')
    add_body_paragraphs(doc, 3, start_idx=4)

    # ===================== CHAPTER 4: Results and Analysis =====================
    add_chapter_title(doc, 'Chapter 4: Results and Analysis')

    add_section_title(doc, '4.1 Experimental Results')
    add_body_paragraphs(doc, 4, start_idx=8)

    add_section_title(doc, '4.2 Comparative Analysis')
    add_body_paragraphs(doc, 4, start_idx=12)

    add_subsection_title(doc, '4.2.1 Performance Benchmarking')
    add_body_paragraphs(doc, 3, start_idx=1)

    add_subsection_title(doc, '4.2.2 Statistical Significance Testing')
    add_body_paragraphs(doc, 3, start_idx=5)

    add_section_title(doc, '4.3 Discussion of Findings')
    add_body_paragraphs(doc, 4, start_idx=9)

    # ===================== CHAPTER 5: Conclusion =====================
    add_chapter_title(doc, 'Chapter 5: Conclusion and Future Work')

    add_section_title(doc, '5.1 Summary of Contributions')
    add_body_paragraphs(doc, 4, start_idx=3)

    # Additional body content for chapter 5
    add_body_paragraphs(doc, 4, start_idx=11)

    # Add extra body paragraphs to ensure ~45 pages
    # Each page holds roughly 25-30 lines with 1.5 spacing at 12pt
    # We need about 30-35 more pages of content
    # Adding filler academic paragraphs throughout

    # Appendix-like additional content to pad pages
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('References')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    references = [
        '[1] Zhang, W., & Patel, R. (2024). "Deep Representation Learning for Cross-Domain Transfer." Journal of Machine Learning Research, 25(3), 1-45.',
        '[2] Nakamura, K., Chen, L., & Okafor, A. (2023). "Scalable Distributed Training for Large Language Models." Proceedings of NeurIPS 2023, pp. 2341-2359.',
        '[3] Rodriguez-Garcia, M. et al. (2024). "Privacy-Preserving Federated Learning with Differential Privacy Guarantees." IEEE Transactions on Pattern Analysis and Machine Intelligence, 46(2), 891-907.',
        '[4] Anderson, J.P., & Williams, S.K. (2023). "Causal Inference in Observational Studies: A Modern Perspective." Statistical Science, 38(4), 567-589.',
        '[5] Liu, X., Gupta, A., & Fernandez, C. (2024). "Attention Mechanisms in Vision Transformers: A Comprehensive Survey." ACM Computing Surveys, 56(1), 1-38.',
        '[6] Thompson, R.A., Park, J.H., & Mueller, K. (2023). "Robust Optimization for Adversarial Machine Learning." ICML 2023 Proceedings, pp. 4521-4539.',
        '[7] Okonkwo, C.E., & Yamamoto, T. (2024). "Graph Neural Networks for Molecular Property Prediction." Nature Machine Intelligence, 6(1), 78-92.',
        '[8] Blackwell, H.M., Santos, P.R., & Kim, D. (2023). "Bayesian Deep Learning: Uncertainty Quantification in Practice." Journal of the Royal Statistical Society, Series B, 85(3), 612-641.',
        '[9] Petrov, I., & Al-Hassan, F. (2024). "Efficient Fine-Tuning of Pre-Trained Models via Low-Rank Adaptation." ICLR 2024 Proceedings, pp. 891-908.',
        '[10] Sullivan, M.J., Tanaka, Y., & Bergstrom, N. (2023). "Reinforcement Learning from Human Feedback: Theory and Applications." Annual Review of Control, Robotics, and Autonomous Systems, 6, 233-261.',
        '[11] Chen, Q., & Adebayo, J. (2024). "Fairness in Machine Learning: Metrics, Methods, and Trade-offs." AI and Ethics Journal, 4(2), 145-172.',
        '[12] Martinez, L.R., Kowalski, P.A., & Huang, W. (2024). "Self-Supervised Visual Representation Learning at Scale." CVPR 2024 Proceedings, pp. 1123-1141.',
        '[13] Johansen, E., & Reeves, D.T. (2023). "Neural Architecture Search: A Comprehensive Survey." Foundations and Trends in Machine Learning, 16(1-2), 1-132.',
        '[14] Krishnan, V., & O\'Brien, K.M. (2024). "Multi-Task Learning for Natural Language Understanding." EMNLP 2024 Proceedings, pp. 3456-3471.',
        '[15] Weber, S., Gupta, N., & Zhao, Y. (2024). "Continual Learning in Dynamic Environments." IEEE Transactions on Neural Networks and Learning Systems, 35(4), 2891-2910.',
        '[16] Patel, A.R., & Johansson, L. (2023). "Energy-Efficient Deep Learning on Edge Devices." Proceedings of MLSys 2023, pp. 789-804.',
        '[17] Foster, K.L., Nakagawa, S., & Dubois, P. (2024). "Generative Models for Scientific Discovery." Science, 383(6679), 142-151.',
        '[18] Hernandez, M., & Sato, T. (2024). "Interpretable Machine Learning for Healthcare Applications." Nature Medicine, 30(2), 234-248.',
        '[19] Clarke, R.J., Liu, Z., & Andersen, P.K. (2023). "Meta-Learning: A Survey of Methods and Applications." Artificial Intelligence Review, 56(8), 9001-9067.',
        '[20] Singh, P., Olsson, C., & Barrett, L.F. (2024). "Multimodal Foundation Models: Challenges and Opportunities." Journal of Artificial Intelligence Research, 79, 543-598.',
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # Add additional appendix content to reach ~45 pages
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('Appendix A: Supplementary Experimental Results')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    # Add extensive appendix content
    appendix_texts = [
        "This appendix provides supplementary experimental results that support the findings presented in Chapter 4. The complete set of ablation studies, hyperparameter sensitivity analyses, and additional baseline comparisons are documented here for reference. All experiments were conducted using NVIDIA A100 GPUs with 80GB memory, running PyTorch 2.1 on Ubuntu 22.04 LTS.",
        "Table A.1 presents the full results of our hyperparameter sweep across learning rates, batch sizes, and regularization coefficients. The optimal configuration was identified through Bayesian optimization using the Expected Improvement acquisition function over 500 iterations. Results are reported as mean and standard deviation across five random seeds.",
        "The convergence analysis in Figure A.1 demonstrates that our proposed method achieves faster convergence compared to all baselines across all three benchmark datasets. The training loss curves show smooth convergence without oscillation, indicating that the adaptive learning rate schedule effectively balances exploration and exploitation during optimization.",
        "Additional ablation studies were conducted to understand the contribution of each architectural component. Removing the cross-attention module resulted in a 3.2% decrease in F1 score on the validation set. Similarly, replacing the learned positional embeddings with sinusoidal embeddings led to a 1.8% performance degradation, suggesting that task-specific position encoding is beneficial.",
        "We also evaluated the robustness of our method under various noise conditions. Gaussian noise with standard deviations ranging from 0.01 to 0.5 was added to the input features. Our method maintained above 90% of its clean performance up to a noise level of 0.2, while the strongest baseline degraded to 78% at the same noise level.",
        "The computational cost analysis reveals that our method requires approximately 2.3x the training time of the vanilla transformer baseline, primarily due to the additional cross-attention computations. However, at inference time, the overhead is minimal (1.1x) because the auxiliary pathways can be pruned after training.",
        "Cross-validation results across ten different data splits confirm the statistical significance of our improvements. A paired t-test yields p-values below 0.001 for all pairwise comparisons against baselines, and the effect sizes (Cohen's d) range from 0.8 to 1.4, indicating large practical significance.",
        "Memory consumption analysis shows that our model requires 12.4 GB of GPU memory during training with a batch size of 32, compared to 8.7 GB for the base transformer. This increase is manageable on modern hardware but may present challenges for deployment on resource-constrained devices.",
        "The feature visualization analysis in Figure A.2 provides qualitative evidence that our model learns more discriminative representations. t-SNE projections of the penultimate layer activations show clearer cluster separation compared to the baseline, particularly for minority classes that are underrepresented in the training data.",
        "We conducted extensive experiments on dataset scaling behavior. When reducing the training set to 10% of its original size, our method retains 85% of its full-data performance, compared to 71% for the strongest baseline. This suggests that the inductive biases introduced by our architectural modifications improve data efficiency.",
        "Transfer learning experiments demonstrate that features learned by our model on the source domain transfer more effectively to related target domains. Fine-tuning on the target task with only 100 labeled examples achieves performance comparable to training from scratch with 10,000 examples.",
        "The error analysis reveals that the remaining failure cases predominantly involve ambiguous instances where even human annotators disagree. Inter-annotator agreement on these difficult examples is only 0.62 (Cohen's kappa), suggesting that our model approaches the ceiling of achievable performance on this benchmark.",
    ]

    for text in appendix_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'

    # Appendix B
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('Appendix B: Mathematical Derivations')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    derivation_texts = [
        "This appendix provides the complete mathematical derivations for the theoretical results presented in Chapter 3. We begin with the proof of Theorem 3.1, which establishes the convergence rate of our proposed optimization algorithm under standard regularity conditions.",
        "Theorem 3.1 (Convergence Rate): Let f be an L-smooth, mu-strongly convex function. Then the iterates x_k generated by Algorithm 1 satisfy E[f(x_k) - f(x*)] <= (1 - mu/L)^k * [f(x_0) - f(x*)], where x* is the unique minimizer of f.",
        "Proof: We proceed by induction on k. The base case k=0 is trivial. For the inductive step, we use the descent lemma for L-smooth functions: f(x_{k+1}) <= f(x_k) + <grad f(x_k), x_{k+1} - x_k> + (L/2)||x_{k+1} - x_k||^2.",
        "Substituting the update rule x_{k+1} = x_k - (1/L) * grad f(x_k) and using the strong convexity inequality, we obtain the desired bound. The complete calculation involves bounding the gradient norm using the PL inequality and telescoping the resulting geometric series.",
        "Lemma B.1 (Generalization Bound): For a hypothesis class H with Rademacher complexity R_n(H), the expected risk satisfies E[L(h)] <= L_emp(h) + 2*R_n(H) + sqrt(ln(2/delta)/(2n)) with probability at least 1-delta over the draw of n training samples.",
        "The proof of Lemma B.1 follows from McDiarmid's inequality combined with the symmetrization technique. The Rademacher complexity term captures the effective capacity of the hypothesis class, while the second term accounts for the finite-sample uncertainty.",
        "Proposition B.2 (Approximation Error): The function class realized by our neural network architecture with depth d and width w can approximate any Lipschitz function on [0,1]^p to within epsilon accuracy using O(epsilon^{-p/s} * log(1/epsilon)) parameters, where s is the smoothness order.",
        "This result extends the classical approximation theory for neural networks to our specific architecture. The proof technique combines covering number arguments with the compositional structure of deep networks, exploiting the fact that composition of Lipschitz functions remains Lipschitz.",
        "The information-theoretic lower bound in Theorem B.3 shows that no estimator can achieve a convergence rate faster than O(n^{-2s/(2s+p)}) in the minimax sense, confirming that our method is rate-optimal up to logarithmic factors.",
        "The derivation of the variational inference objective in Section 3.4 begins with the standard evidence lower bound (ELBO): log p(x) >= E_q[log p(x|z)] - KL(q(z|x) || p(z)), where q(z|x) is the approximate posterior and p(z) is the prior.",
        "By choosing a flexible normalizing flow for q(z|x), we can tighten the ELBO significantly. The change of variables formula yields: log q_K(z_K) = log q_0(z_0) - sum_{k=1}^K log |det J_k|, where J_k is the Jacobian of the k-th flow transformation.",
        "Finally, the gradient estimator in Algorithm 2 combines the pathwise gradient with a control variate to reduce variance. The optimal control variate coefficient is derived by minimizing the trace of the gradient covariance matrix, yielding a closed-form expression involving the Fisher information.",
    ]

    for text in derivation_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'

    # Appendix C - more content to fill pages
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('Appendix C: Implementation Details')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    impl_texts = [
        "This appendix provides comprehensive implementation details to facilitate reproducibility of our experimental results. All source code is available at the project repository, and we provide Docker containers with pre-configured environments for each experiment.",
        "The base model architecture consists of 12 transformer layers with 768-dimensional hidden states, 12 attention heads, and a vocabulary size of 50,257 tokens. The model was initialized with pre-trained weights from a publicly available checkpoint and fine-tuned on our task-specific datasets.",
        "Training was conducted using the AdamW optimizer with an initial learning rate of 2e-5, weight decay of 0.01, and a linear warmup schedule over the first 10% of training steps followed by cosine decay to zero. Gradient clipping with a maximum norm of 1.0 was applied throughout training.",
        "Data preprocessing involved tokenization using the SentencePiece algorithm with a vocabulary of 32,000 subword units. Input sequences were truncated or padded to a maximum length of 512 tokens, and attention masks were applied to exclude padding positions from the attention computation.",
        "Hyperparameter tuning was performed using a combination of grid search and Bayesian optimization. The grid search explored learning rates in {1e-5, 2e-5, 5e-5, 1e-4}, batch sizes in {16, 32, 64}, and dropout rates in {0.1, 0.2, 0.3}. Bayesian optimization was then used to refine the best configuration.",
        "The evaluation pipeline consists of three stages: (1) model inference on the test set, (2) post-processing and decoding of model outputs, and (3) metric computation. We report precision, recall, F1 score, and accuracy for classification tasks, and BLEU, ROUGE-L, and BERTScore for generation tasks.",
        "For the distributed training setup, we used PyTorch's DistributedDataParallel (DDP) module with NCCL backend across 8 NVIDIA A100 GPUs. Gradient synchronization was performed after each backward pass, and we used mixed-precision training (FP16) with dynamic loss scaling to accelerate training.",
        "The inference pipeline was optimized for production deployment using ONNX Runtime with TensorRT backend. Model quantization to INT8 was applied using post-training quantization with a calibration dataset of 1,000 samples. This reduced the model size by 4x and improved inference throughput by 3.2x.",
        "Monitoring and logging were handled through Weights & Biases (wandb), with automatic tracking of loss curves, gradient norms, learning rate schedules, and evaluation metrics. All experiment configurations were versioned using hydra-zen, enabling exact reproduction of any experimental run.",
        "Error handling and fault tolerance were implemented through periodic checkpointing (every 1000 steps) and automatic resume from the latest checkpoint. The training script also includes deadlock detection and automatic restart capabilities for distributed training scenarios.",
        "The dataset processing pipeline was implemented using Apache Beam for distributed data processing, enabling efficient handling of datasets exceeding 100GB. The pipeline includes deduplication, quality filtering, and privacy-preserving data transformations.",
        "Unit tests and integration tests were developed for all critical components of the codebase. The test suite achieves 92% code coverage and includes property-based testing using Hypothesis for numerical computations. Continuous integration was configured through GitHub Actions with automated testing on each pull request.",
    ]

    for text in impl_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')


def launch_gui(command, delay_sec=1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


create_initial()

# Open in LibreOffice Writer
launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')
