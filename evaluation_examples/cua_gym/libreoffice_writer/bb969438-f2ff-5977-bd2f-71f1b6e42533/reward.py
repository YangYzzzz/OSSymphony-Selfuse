"""
Reward Script: Set up chapter-based figure numbering in thesis
Task ID: writer_acad_055
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): All figure captions use "Figure X.Y:" format (chapter.seq)
  Component 2 (0.3): Chapter prefix in each caption matches the actual chapter number
  Component 3 (0.3): In-text figure references also use chapter-based numbering
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_055'


def persist_app_state(domain):
    """Try to save any unsaved edits in LibreOffice Writer."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_writer", "libreoffice_calc", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that figure numbering has been changed from sequential (Figure 1, 2, ...)
    to chapter-based (Figure 1.1, 1.2, 2.1, ...) throughout the document.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Build chapter map: paragraph index -> chapter number
    # and collect figure caption paragraphs and reference paragraphs
    chapter_num = 0
    para_chapter = {}  # para index -> chapter number
    caption_paras = []  # (para_index, para_text, chapter_num)
    ref_paras = []      # (para_index, para_text, chapter_num)

    chapter_prefix_pattern = re.compile(r'^Figure\s+(\d+)\.(\d+)\s*:')
    sequential_pattern = re.compile(r'^Figure\s+(\d+)\s*:')
    # Pattern for in-text references like "Figure 1.1" or "Figures 2.2 and 2.3"
    ref_chapter_pattern = re.compile(r'Figure[s]?\s+(\d+\.\d+)')

    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ''
        text = para.text.strip()

        if style == 'Heading 1':
            chapter_num += 1

        para_chapter[i] = chapter_num

        # Identify figure caption paragraphs (start with "Figure X.Y:" or "Figure N:")
        if chapter_prefix_pattern.match(text) or sequential_pattern.match(text):
            caption_paras.append((i, text, chapter_num))
        elif ref_chapter_pattern.search(text) and not chapter_prefix_pattern.match(text):
            # This is a paragraph with in-text figure references (but not a caption itself)
            ref_paras.append((i, text, chapter_num))

    if not caption_paras:
        print("FAIL: No figure caption paragraphs found in the document")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Found {len(caption_paras)} figure caption paragraphs")
    print(f"INFO: Found {len(ref_paras)} paragraphs with in-text figure references")
    print(f"INFO: Found {chapter_num} chapters (Heading 1)")

    # Component 1: All figure captions use "Figure X.Y:" format (0.4 points)
    # This checks that captions have chapter-based numbering, not sequential.
    try:
        total_captions = len(caption_paras)
        chapter_format_count = 0
        for idx, text, ch in caption_paras:
            m = chapter_prefix_pattern.match(text)
            if m:
                chapter_format_count += 1
                print(f"  PASS: Para {idx} uses chapter format: Figure {m.group(1)}.{m.group(2)}")
            else:
                print(f"  FAIL: Para {idx} does NOT use chapter format: {text[:60]}")

        if total_captions > 0 and chapter_format_count == total_captions:
            print(f"PASS: Component 1 -- All {total_captions} captions use Figure X.Y format (0.4 pts)")
            total_score += 0.4
        elif total_captions > 0 and chapter_format_count > 0:
            partial = 0.4 * (chapter_format_count / total_captions)
            print(f"PARTIAL: Component 1 -- {chapter_format_count}/{total_captions} captions use chapter format ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 -- No captions use chapter-based numbering format")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Chapter prefix matches actual chapter (0.3 points)
    # e.g., Figure 2.1 must appear inside Chapter 2 (Heading 1 "Related Work")
    try:
        correct_prefix_count = 0
        checked_count = 0
        for idx, text, actual_chapter in caption_paras:
            m = chapter_prefix_pattern.match(text)
            if m:
                claimed_chapter = int(m.group(1))
                checked_count += 1
                if claimed_chapter == actual_chapter:
                    correct_prefix_count += 1
                    print(f"  PASS: Figure {m.group(1)}.{m.group(2)} correctly in chapter {actual_chapter}")
                else:
                    print(f"  FAIL: Figure {m.group(1)}.{m.group(2)} claims chapter {claimed_chapter} but is in chapter {actual_chapter}")

        if checked_count > 0 and correct_prefix_count == checked_count:
            print(f"PASS: Component 2 -- All {checked_count} captions have correct chapter prefix (0.3 pts)")
            total_score += 0.3
        elif checked_count > 0 and correct_prefix_count > 0:
            partial = 0.3 * (correct_prefix_count / checked_count)
            print(f"PARTIAL: Component 2 -- {correct_prefix_count}/{checked_count} have correct prefix ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- No captions have correct chapter prefix (or none in chapter format)")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: In-text references use chapter-based numbering (0.3 points)
    # Check that references like "Figure 1.1" appear in the text (not "Figure 1")
    try:
        # Collect all in-text figure references across entire document (excluding captions)
        all_refs = []
        sequential_refs = []
        chapter_refs = []

        sequential_ref_pattern = re.compile(r'(?<!\d\.)Figure[s]?\s+(\d+)(?!\.\d)')
        chapter_ref_inline = re.compile(r'Figure[s]?\s+\d+\.\d+')

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            # Skip caption paragraphs and empty paragraphs
            if chapter_prefix_pattern.match(text) or sequential_pattern.match(text):
                continue
            if not text:
                continue

            # Find chapter-based references
            ch_matches = chapter_ref_inline.findall(text)
            chapter_refs.extend(ch_matches)

            # Find sequential references (bad - means not updated)
            # Must be careful: "Figure 1.1" should not match as sequential
            seq_matches = sequential_ref_pattern.findall(text)
            # Filter out numbers that are part of chapter refs already matched
            for sm in seq_matches:
                # Check if this is actually part of "Figure X.Y"
                # The negative lookbehind/lookahead should handle this, but double check
                if f"Figure {sm}." not in text and f"Figures {sm}." not in text:
                    sequential_refs.append(sm)

        total_refs = len(chapter_refs) + len(sequential_refs)
        print(f"  INFO: Found {len(chapter_refs)} chapter-based refs, {len(sequential_refs)} sequential refs")

        if total_refs > 0 and len(chapter_refs) > 0 and len(sequential_refs) == 0:
            print(f"PASS: Component 3 -- All in-text references use chapter-based numbering (0.3 pts)")
            total_score += 0.3
        elif total_refs > 0 and len(chapter_refs) > 0:
            ratio = len(chapter_refs) / total_refs
            partial = 0.3 * ratio
            print(f"PARTIAL: Component 3 -- {len(chapter_refs)}/{total_refs} refs are chapter-based ({partial:.2f} pts)")
            total_score += partial
        elif len(chapter_refs) == 0 and total_refs == 0:
            # No in-text references at all - this is unusual but not a scoring failure
            # Check if the golden has refs; if there are none, give benefit of doubt
            print(f"INFO: Component 3 -- No in-text figure references found; skipping")
            # Don't award points - the task says numbering should be updated throughout
            print(f"FAIL: Component 3 -- Expected in-text references with chapter numbering")
        else:
            print(f"FAIL: Component 3 -- No chapter-based in-text references found ({len(sequential_refs)} sequential refs)")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
