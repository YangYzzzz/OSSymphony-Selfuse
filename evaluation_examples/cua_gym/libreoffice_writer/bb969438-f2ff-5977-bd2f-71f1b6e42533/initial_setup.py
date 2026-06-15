"""
Initial Setup: Set up figure numbering per chapter in thesis
Task ID: writer_acad_055
Domain: libreoffice_writer

Creates a thesis document with 3 chapters and 8 figures.
Figures are numbered sequentially (Figure 1 through Figure 8).
Chapter headings use Heading 1 style with outline numbering.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_055'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_chapter_numbering(doc):
    """Add outline numbering to Heading 1 style so chapters are numbered 1, 2, 3..."""
    # Create abstractNum for chapter numbering
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part.numbering_definitions._numbering

    # Define abstract numbering for chapter headings
    abstract_num_xml = (
        '<w:abstractNum w:abstractNumId="100" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '  <w:multiLevelType w:val="multilevel"/>'
        '  <w:lvl w:ilvl="0">'
        '    <w:start w:val="1"/>'
        '    <w:numFmt w:val="decimal"/>'
        '    <w:pStyle w:val="Heading1"/>'
        '    <w:lvlText w:val="Chapter %1"/>'
        '    <w:lvlJc w:val="left"/>'
        '  </w:lvl>'
        '</w:abstractNum>'
    )
    from lxml import etree
    abstract_num = etree.fromstring(abstract_num_xml)
    numbering_elm.insert(0, abstract_num)

    # Create num referencing abstractNum
    num_xml = (
        '<w:num w:numId="100" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '  <w:abstractNumId w:val="100"/>'
        '</w:num>'
    )
    num_elm = etree.fromstring(num_xml)
    numbering_elm.append(num_elm)

    # Link Heading 1 style to numbering
    styles = doc.styles
    heading1 = styles['Heading 1']
    heading1_element = heading1.element
    pPr = heading1_element.find(qn('w:pPr'))
    if pPr is None:
        pPr = heading1_element.makeelement(qn('w:pPr'), {})
        heading1_element.insert(0, pPr)
    numPr = pPr.makeelement(qn('w:numPr'), {})
    ilvl = numPr.makeelement(qn('w:ilvl'), {qn('w:val'): '0'})
    numId = numPr.makeelement(qn('w:numId'), {qn('w:val'): '100'})
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def add_figure_caption(doc, fig_number, caption_text):
    """Add a figure caption with sequential numbering: Figure N: caption_text"""
    cap_para = doc.add_paragraph()
    cap_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cap_para.paragraph_format.space_before = Pt(6)
    cap_para.paragraph_format.space_after = Pt(12)

    run = cap_para.add_run(f"Figure {fig_number}: {caption_text}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    return cap_para


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set margins
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title page
    title = doc.add_heading('Adaptive Neural Architecture Search for Edge Computing Applications', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy in Computer Science")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run("\n\nby\nElena Vasquez-Rodriguez\n\nDepartment of Computer Science\nStanford University\nMarch 2025")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_page_break()

    # Add chapter numbering to Heading 1 style
    add_chapter_numbering(doc)

    # ---- CHAPTER 1: Introduction ----
    doc.add_heading('Introduction', level=1)

    doc.add_paragraph(
        "The rapid proliferation of edge computing devices has created an urgent need for neural "
        "network architectures that balance computational efficiency with predictive accuracy. "
        "Traditional deep learning models, designed for cloud-based inference, often exceed the "
        "memory and power constraints of embedded systems such as autonomous drones, wearable "
        "health monitors, and industrial IoT sensors."
    )

    doc.add_paragraph(
        "Neural Architecture Search (NAS) has emerged as a promising approach to automatically "
        "discover architectures optimized for specific hardware targets. However, existing NAS "
        "methods face two critical limitations: prohibitive search costs and poor transferability "
        "across hardware platforms. This thesis addresses both challenges through a novel adaptive "
        "search framework."
    )

    # Figure 1 - placeholder image description as paragraph
    fig1_para = doc.add_paragraph()
    fig1_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig1_para.add_run("[Diagram showing the growth of edge devices from 2018-2024]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_figure_caption(doc, 1, "Global edge computing device deployment trends (2018-2024)")

    doc.add_paragraph(
        "As illustrated in Figure 1, the number of deployed edge devices has grown exponentially, "
        "reaching an estimated 15.1 billion units by the end of 2024. Each device category presents "
        "unique computational constraints that necessitate specialized architectures. The heterogeneity "
        "of deployment targets makes manual architecture design increasingly impractical."
    )

    # Figure 2
    fig2_para = doc.add_paragraph()
    fig2_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig2_para.add_run("[Architecture comparison chart: Cloud vs. Edge model parameters]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_figure_caption(doc, 2, "Parameter count comparison between cloud and edge-optimized architectures")

    doc.add_paragraph(
        "Figure 2 demonstrates the significant gap between cloud-scale and edge-scale architectures. "
        "While transformer-based models such as GPT-4 operate with hundreds of billions of parameters, "
        "edge deployment targets typically require models below 10 million parameters with inference "
        "latency under 50 milliseconds."
    )

    # Figure 3
    fig3_para = doc.add_paragraph()
    fig3_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig3_para.add_run("[Flowchart of the proposed adaptive NAS framework overview]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_figure_caption(doc, 3, "Overview of the proposed Adaptive Neural Architecture Search (ANAS) framework")

    doc.add_paragraph(
        "The proposed ANAS framework, depicted in Figure 3, introduces a three-stage pipeline: "
        "hardware profiling, architecture generation via differentiable search, and deployment-aware "
        "fine-tuning. This integrated approach reduces search cost by 73% compared to prior methods."
    )

    doc.add_page_break()

    # ---- CHAPTER 2: Related Work ----
    doc.add_heading('Related Work', level=1)

    doc.add_paragraph(
        "This chapter surveys the foundational literature across three interconnected domains: "
        "neural architecture search, model compression techniques, and hardware-aware optimization. "
        "We identify critical gaps in existing approaches that motivate our proposed framework."
    )

    doc.add_heading('Neural Architecture Search', level=2)

    doc.add_paragraph(
        "Early NAS methods relied on reinforcement learning controllers to sample architectures "
        "from a discrete search space. Zoph and Le (2017) demonstrated that RL-based search could "
        "discover architectures competitive with hand-designed networks on CIFAR-10, but required "
        "over 22,400 GPU hours. Subsequent work by Pham et al. (2018) introduced weight sharing "
        "through ENAS, reducing search cost by three orders of magnitude."
    )

    # Figure 4
    fig4_para = doc.add_paragraph()
    fig4_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig4_para.add_run("[Timeline showing evolution of NAS methods from 2017 to 2024]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_figure_caption(doc, 4, "Evolution of NAS methodologies and their computational requirements")

    doc.add_paragraph(
        "Figure 4 traces the progression from early resource-intensive approaches to modern "
        "efficient search strategies. The trend toward differentiable methods, initiated by DARTS "
        "(Liu et al., 2019), significantly democratized architecture search by enabling gradient-based "
        "optimization over continuous relaxations of the search space."
    )

    doc.add_heading('Model Compression and Pruning', level=2)

    doc.add_paragraph(
        "Complementary to NAS, model compression techniques seek to reduce the computational "
        "footprint of pre-existing architectures. Han et al. (2016) introduced magnitude-based "
        "weight pruning, achieving 90% sparsity on AlexNet with minimal accuracy degradation. "
        "More recent work by Frankle and Carlin (2019) proposed the Lottery Ticket Hypothesis, "
        "suggesting that dense networks contain sparse sub-networks capable of matching full "
        "network performance when trained in isolation."
    )

    # Figure 5
    fig5_para = doc.add_paragraph()
    fig5_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig5_para.add_run("[Bar chart comparing compression ratios across pruning methods]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_figure_caption(doc, 5, "Compression ratio vs. accuracy trade-off for leading pruning methods on ImageNet")

    # Figure 6
    fig6_para = doc.add_paragraph()
    fig6_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig6_para.add_run("[Scatter plot of latency vs. accuracy for compressed models]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_figure_caption(doc, 6, "Latency-accuracy Pareto frontier for compressed architectures on Cortex-M7")

    doc.add_paragraph(
        "Figures 5 and 6 reveal a critical insight: while compression methods effectively reduce "
        "model size, the resulting latency profiles are highly dependent on the target hardware. "
        "A model pruned for GPU efficiency may exhibit suboptimal performance on microcontrollers "
        "due to irregular memory access patterns."
    )

    doc.add_page_break()

    # ---- CHAPTER 3: Methodology ----
    doc.add_heading('Methodology', level=1)

    doc.add_paragraph(
        "This chapter presents the technical details of the Adaptive Neural Architecture Search "
        "(ANAS) framework. We formalize the multi-objective search problem, describe the hardware "
        "profiling mechanism, and detail the differentiable search algorithm with deployment-aware "
        "constraints."
    )

    doc.add_heading('Problem Formulation', level=2)

    doc.add_paragraph(
        "Let A denote the architecture search space and H represent a set of target hardware "
        "platforms. For each candidate architecture a in A and hardware target h in H, we define "
        "three objective functions: accuracy f_acc(a), latency f_lat(a, h), and energy consumption "
        "f_eng(a, h). The multi-objective optimization problem seeks to find the Pareto-optimal set "
        "of architectures that simultaneously minimize latency and energy while maximizing accuracy."
    )

    # Figure 7
    fig7_para = doc.add_paragraph()
    fig7_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig7_para.add_run("[Mathematical formulation and search space visualization]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_figure_caption(doc, 7, "Search space structure with hardware-aware constraint boundaries")

    doc.add_heading('Hardware Profiling Module', level=2)

    doc.add_paragraph(
        "The hardware profiling module constructs latency and energy lookup tables for each "
        "operator type on each target platform. Rather than relying on FLOPs as a proxy metric, "
        "we directly measure execution time through automated benchmarking. Each operation in the "
        "search space is profiled across batch sizes ranging from 1 to 32, capturing the "
        "non-linear relationship between batch size and throughput."
    )

    # Figure 8
    fig8_para = doc.add_paragraph()
    fig8_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig8_para.add_run("[Heatmap of operator latencies across different hardware platforms]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_figure_caption(doc, 8, "Operator latency profiles across five target hardware platforms")

    doc.add_paragraph(
        "Figure 8 highlights the dramatic variation in operator performance across platforms. "
        "Depthwise separable convolutions, for instance, are 3.2x faster than standard convolutions "
        "on the Jetson Nano but only 1.4x faster on the Cortex-M7, underscoring the need for "
        "platform-specific architecture decisions."
    )

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
