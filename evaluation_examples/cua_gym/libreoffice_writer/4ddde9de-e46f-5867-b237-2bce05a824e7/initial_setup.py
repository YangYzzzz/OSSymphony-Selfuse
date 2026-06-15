"""
Initial Setup: Create a 20-page Master Services Agreement with no headers configured.
Task ID: writer_legal_057
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_057'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_para(doc, text, level=1):
    """Add a heading with appropriate formatting."""
    if level == 0:
        h = doc.add_heading(text, level=0)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif level == 1:
        h = doc.add_heading(text, level=1)
    elif level == 2:
        h = doc.add_heading(text, level=2)
    else:
        h = doc.add_heading(text, level=3)
    return h


def add_body(doc, text, bold=False, space_after=Pt(6)):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


def create_initial():
    doc = Document()

    # Page setup - standard letter
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ---- PAGE 1: Title Page ----
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_heading('MASTER SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(24)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Between')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('ACME CORPORATION')
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('("Service Provider")')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('and')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('BETA INCORPORATED')
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('("Client")')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Effective Date: January 15, 2025')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Agreement No.: MSA-2025-0417')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Page break after title page
    doc.add_page_break()

    # ---- RECITALS ----
    add_heading_para(doc, 'RECITALS', level=1)

    add_body(doc, 'WHEREAS, Acme Corporation ("Acme") is a Delaware corporation engaged in the business of providing professional technology consulting, software development, systems integration, and managed IT services to enterprise clients across multiple industry verticals;')

    add_body(doc, 'WHEREAS, Beta Incorporated ("Beta") is a California corporation that desires to engage the services of Acme to provide certain professional technology services, software development support, and related consulting as described herein and in subsequent Statements of Work;')

    add_body(doc, 'WHEREAS, both parties wish to establish the terms and conditions under which Acme will provide such services, and the rights and obligations of each party with respect thereto;')

    add_body(doc, 'NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth herein, and for other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, the parties agree as follows:')

    doc.add_paragraph()

    # ---- ARTICLE 1: DEFINITIONS ----
    add_heading_para(doc, 'ARTICLE 1: DEFINITIONS AND INTERPRETATION', level=1)

    definitions = [
        ('"Affiliate"', 'means any entity that directly or indirectly controls, is controlled by, or is under common control with a party, where "control" means the possession, directly or indirectly, of the power to direct or cause the direction of the management and policies of an entity, whether through ownership of voting securities, by contract, or otherwise.'),
        ('"Agreement"', 'means this Master Services Agreement, together with all Exhibits, Schedules, and Statements of Work attached hereto or incorporated herein by reference, as amended from time to time in accordance with Section 18.5.'),
        ('"Change Order"', 'means a written document executed by both parties that modifies a Statement of Work, including any changes to scope, timeline, deliverables, fees, or resource allocation.'),
        ('"Confidential Information"', 'means all non-public information disclosed by one party to the other, whether orally, in writing, or by inspection, that is designated as confidential or that reasonably should be understood to be confidential given the nature of the information and the circumstances of disclosure. Confidential Information includes, but is not limited to, trade secrets, business plans, financial data, customer lists, technical specifications, source code, algorithms, and proprietary methodologies.'),
        ('"Deliverable"', 'means any tangible or intangible work product, including software, documentation, reports, designs, specifications, or other materials, that Acme is required to deliver to Beta under a Statement of Work.'),
        ('"Effective Date"', 'means January 15, 2025, being the date first written above.'),
        ('"Fees"', 'means all compensation payable by Beta to Acme for Services rendered under this Agreement and any Statement of Work, as detailed in Article 7 and the applicable SOW.'),
        ('"Force Majeure Event"', 'means any event beyond the reasonable control of the affected party, including but not limited to acts of God, natural disasters, epidemics, pandemics, war, terrorism, civil unrest, government actions, embargoes, labor disputes, power failures, telecommunications failures, cyberattacks, or supply chain disruptions.'),
        ('"Intellectual Property"', 'means all patents, copyrights, trademarks, trade secrets, know-how, moral rights, and all other intellectual property rights, whether registered or unregistered, and all applications and rights to apply for any of the foregoing, anywhere in the world.'),
        ('"Services"', 'means the professional technology consulting, software development, systems integration, project management, quality assurance, and other services to be provided by Acme to Beta under this Agreement, as further described in applicable Statements of Work.'),
        ('"Statement of Work" or "SOW"', 'means a document, substantially in the form attached as Exhibit A, that describes specific Services to be performed, Deliverables to be provided, timelines, milestones, acceptance criteria, and Fees for a particular engagement under this Agreement.'),
    ]

    for term, defn in definitions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        r1 = p.add_run(term + ' ')
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(11)
        r2 = p.add_run(defn)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(11)

    # ---- ARTICLE 2: SCOPE OF SERVICES ----
    add_heading_para(doc, 'ARTICLE 2: SCOPE OF SERVICES', level=1)

    add_body(doc, '2.1 General Scope. Acme agrees to provide the Services described in this Agreement and in each Statement of Work duly executed by both parties. The Services shall be performed in a professional and workmanlike manner consistent with generally accepted industry standards and practices applicable to the type of Services being provided.')

    add_body(doc, '2.2 Statements of Work. The specific Services to be performed by Acme shall be described in individual Statements of Work, each of which shall be executed by authorized representatives of both parties and shall be incorporated into and governed by the terms of this Agreement. Each SOW shall specify, at a minimum: (a) a description of the Services to be performed; (b) the Deliverables to be provided; (c) the project timeline and milestones; (d) the acceptance criteria for each Deliverable; (e) the Fees and payment schedule; (f) the Acme personnel assigned to the project; and (g) any specific terms and conditions applicable to that engagement.')

    add_body(doc, '2.3 Priority and Conflicts. In the event of any conflict between the terms of this Agreement and the terms of any Statement of Work, the terms of this Agreement shall control unless the SOW expressly states that it is intended to modify a specific provision of this Agreement and is signed by authorized officers of both parties.')

    add_body(doc, '2.4 Exclusivity. Nothing in this Agreement shall be construed to grant Beta any exclusive right to the Services of Acme. Acme retains the right to provide similar services to other clients, provided that such engagements do not conflict with Acme\'s obligations under this Agreement or any applicable SOW, and provided further that Acme complies with the confidentiality obligations set forth in Article 9.')

    add_body(doc, '2.5 Subcontracting. Acme may, with the prior written consent of Beta (such consent not to be unreasonably withheld, conditioned, or delayed), engage subcontractors to perform portions of the Services. Acme shall remain fully responsible for the performance of any subcontractor and shall ensure that each subcontractor is bound by obligations at least as protective as those set forth in this Agreement, including without limitation the confidentiality and intellectual property provisions.')

    # ---- ARTICLE 3: PERSONNEL AND RESOURCES ----
    add_heading_para(doc, 'ARTICLE 3: PERSONNEL AND RESOURCES', level=1)

    add_body(doc, '3.1 Key Personnel. Acme shall assign qualified personnel to perform the Services. Key Personnel for each engagement shall be identified in the applicable SOW. Acme shall not remove or reassign Key Personnel without the prior written consent of Beta, except in cases of voluntary resignation, termination for cause, or medical leave, in which case Acme shall promptly notify Beta and provide a qualified replacement.')

    add_body(doc, '3.2 Qualifications. All personnel assigned to perform Services shall possess the skills, experience, and qualifications necessary to perform the Services to which they are assigned. Acme represents and warrants that all personnel have been properly vetted, including background checks where required by applicable law or industry standards.')

    add_body(doc, '3.3 Beta Resources. Beta shall provide Acme with reasonable access to Beta\'s facilities, systems, equipment, and personnel as necessary for Acme to perform the Services. Beta shall designate a project manager or liaison to coordinate with Acme and shall ensure timely provision of information, decisions, and approvals reasonably required for Acme to perform its obligations under each SOW.')

    add_body(doc, '3.4 Working Conditions. All Acme personnel working on Beta\'s premises shall comply with Beta\'s workplace policies, security procedures, and code of conduct. Beta shall provide a safe working environment for Acme personnel in accordance with applicable occupational health and safety regulations.')

    # ---- ARTICLE 4: PROJECT MANAGEMENT ----
    add_heading_para(doc, 'ARTICLE 4: PROJECT MANAGEMENT AND GOVERNANCE', level=1)

    add_body(doc, '4.1 Project Governance. Each SOW shall establish a governance structure appropriate to the scope and complexity of the engagement. At a minimum, each engagement shall include: (a) a designated Project Manager from each party; (b) regular status meetings at intervals specified in the SOW; (c) written status reports from Acme detailing progress against milestones, risks, issues, and resource utilization; and (d) an escalation procedure for resolving disputes or issues that cannot be resolved at the project level.')

    add_body(doc, '4.2 Change Management. Either party may propose changes to an SOW through a written Change Order request. All Change Orders must be reviewed, negotiated in good faith, and executed by authorized representatives of both parties before implementation. No changes to scope, timeline, or Fees shall be effective until documented in a fully executed Change Order.')

    add_body(doc, '4.3 Steering Committee. For engagements exceeding $500,000 in total Fees or 12 months in duration, the parties shall establish a Steering Committee consisting of senior representatives from each party. The Steering Committee shall meet no less than quarterly to review overall program status, strategic alignment, and resolution of escalated issues.')

    add_body(doc, '4.4 Risk Management. Acme shall maintain a project risk register for each engagement and shall promptly notify Beta of any material risks that could adversely affect the delivery of Services, the quality of Deliverables, or the project timeline. Both parties shall cooperate in good faith to develop and implement risk mitigation strategies.')

    # ---- ARTICLE 5: DELIVERABLES AND ACCEPTANCE ----
    add_heading_para(doc, 'ARTICLE 5: DELIVERABLES AND ACCEPTANCE', level=1)

    add_body(doc, '5.1 Delivery. Acme shall deliver all Deliverables in accordance with the schedule and specifications set forth in the applicable SOW. Acme shall use commercially reasonable efforts to meet all deadlines and milestones, and shall promptly notify Beta if Acme reasonably believes that any deadline or milestone may not be met, together with a proposed remediation plan.')

    add_body(doc, '5.2 Acceptance Testing. Upon delivery of each Deliverable, Beta shall have a period of fifteen (15) business days (or such other period as may be specified in the applicable SOW) to review and test the Deliverable against the acceptance criteria specified in the SOW (the "Acceptance Period"). Beta shall notify Acme in writing of acceptance or rejection of the Deliverable within the Acceptance Period.')

    add_body(doc, '5.3 Rejection and Cure. If Beta reasonably rejects a Deliverable, Beta\'s notice of rejection shall include a detailed description of the deficiencies. Acme shall correct the identified deficiencies and re-deliver the Deliverable within ten (10) business days. Beta shall then have an additional ten (10) business day Acceptance Period to review the corrected Deliverable. If the Deliverable fails acceptance testing a second time, Beta may, at its option: (a) grant Acme an additional cure period; (b) reduce the Fees proportionally; or (c) terminate the applicable SOW in accordance with Article 14.')

    add_body(doc, '5.4 Deemed Acceptance. If Beta fails to provide written notice of acceptance or rejection within the Acceptance Period, the Deliverable shall be deemed accepted. Notwithstanding the foregoing, deemed acceptance shall not relieve Acme of its warranty obligations under Article 12.')

    # ---- ARTICLE 6: TERM AND TERMINATION ----
    add_heading_para(doc, 'ARTICLE 6: TERM AND RENEWAL', level=1)

    add_body(doc, '6.1 Initial Term. This Agreement shall commence on the Effective Date and shall continue for a period of three (3) years thereafter (the "Initial Term"), unless earlier terminated in accordance with Article 14.')

    add_body(doc, '6.2 Renewal. This Agreement shall automatically renew for successive one (1) year periods (each a "Renewal Term") unless either party provides written notice of non-renewal at least ninety (90) days prior to the expiration of the then-current term. The Initial Term and any Renewal Terms are collectively referred to as the "Term."')

    add_body(doc, '6.3 SOW Survival. The expiration or termination of this Agreement shall not affect any SOW that is then in effect, unless specifically provided otherwise. Each outstanding SOW shall continue in accordance with its terms until completed, terminated, or expired in accordance with its own provisions, subject to the continuing applicability of Articles 8, 9, 10, 11, 15, 16, and 17 of this Agreement.')

    # ---- ARTICLE 7: FEES AND PAYMENT ----
    add_heading_para(doc, 'ARTICLE 7: FEES AND PAYMENT', level=1)

    add_body(doc, '7.1 Fees. Beta shall pay Acme the Fees specified in each SOW for Services rendered. Unless otherwise specified in the applicable SOW, Fees shall be calculated on a time-and-materials basis at the hourly rates set forth in Exhibit B (Rate Card), as updated from time to time in accordance with Section 7.6.')

    add_body(doc, '7.2 Invoicing. Acme shall submit invoices to Beta on a monthly basis (or as otherwise specified in the applicable SOW) detailing the Services performed, hours expended by each resource, expenses incurred, and the total amount due. Each invoice shall be accompanied by reasonable supporting documentation.')

    add_body(doc, '7.3 Payment Terms. Beta shall pay each undisputed invoice within thirty (30) days of receipt. Late payments shall bear interest at the lesser of one and one-half percent (1.5%) per month or the maximum rate permitted by applicable law. Beta shall not withhold or set off any amounts due under this Agreement except as expressly permitted herein.')

    add_body(doc, '7.4 Expenses. Acme shall be entitled to reimbursement for reasonable and necessary out-of-pocket expenses incurred in connection with the performance of Services, provided that: (a) expenses exceeding $500 per item require Beta\'s prior written approval; (b) all expenses are documented with original receipts; and (c) travel expenses comply with Beta\'s travel policy as provided to Acme in writing.')

    add_body(doc, '7.5 Taxes. All Fees are exclusive of applicable taxes. Beta shall be responsible for all sales, use, value-added, and similar taxes arising from the transactions contemplated by this Agreement, excluding taxes based on Acme\'s net income. If Beta is required by law to withhold taxes from payments to Acme, Beta shall provide Acme with documentation sufficient to enable Acme to claim applicable tax credits or refunds.')

    add_body(doc, '7.6 Rate Adjustments. Acme may adjust the rates set forth in Exhibit B no more than once per calendar year, effective upon at least sixty (60) days\' prior written notice to Beta. Rate increases shall not exceed the greater of: (a) five percent (5%) of the then-current rates; or (b) the percentage increase in the Consumer Price Index (All Urban Consumers) for the preceding twelve-month period.')

    # ---- ARTICLE 8: INTELLECTUAL PROPERTY ----
    add_heading_para(doc, 'ARTICLE 8: INTELLECTUAL PROPERTY RIGHTS', level=1)

    add_body(doc, '8.1 Pre-Existing IP. Each party shall retain all right, title, and interest in and to its pre-existing Intellectual Property. Neither party grants the other any rights in its pre-existing IP except as expressly provided in this Agreement or any SOW.')

    add_body(doc, '8.2 Work Product. Unless otherwise specified in the applicable SOW, all Deliverables and work product created by Acme specifically for Beta in the performance of Services under this Agreement ("Work Product") shall be considered works made for hire to the extent permitted by applicable law. To the extent any Work Product does not qualify as a work made for hire, Acme hereby irrevocably assigns to Beta all right, title, and interest in and to such Work Product, including all Intellectual Property rights therein.')

    add_body(doc, '8.3 Acme Tools and Methodologies. Notwithstanding Section 8.2, Acme shall retain all right, title, and interest in and to: (a) pre-existing tools, frameworks, libraries, methodologies, and know-how developed by Acme independently of this Agreement; (b) any improvements or enhancements to such pre-existing materials made during the performance of Services; and (c) generic utilities, templates, and development tools created during the engagement that are not specific to Beta\'s business. Acme hereby grants Beta a non-exclusive, perpetual, irrevocable, royalty-free license to use any such Acme materials embedded in the Deliverables.')

    add_body(doc, '8.4 Open Source. Acme shall not incorporate any open source software into any Deliverable without Beta\'s prior written approval. Any request for approval shall identify the open source component, its license terms, and any potential impact on Beta\'s rights in the Deliverable. Acme shall maintain a register of all open source components incorporated into Deliverables and shall provide such register to Beta upon request.')

    # ---- ARTICLE 9: CONFIDENTIALITY ----
    add_heading_para(doc, 'ARTICLE 9: CONFIDENTIALITY', level=1)

    add_body(doc, '9.1 Obligations. Each party (the "Receiving Party") shall: (a) hold the Confidential Information of the other party (the "Disclosing Party") in strict confidence; (b) not disclose such Confidential Information to any third party without the prior written consent of the Disclosing Party, except to employees, agents, and subcontractors who have a need to know and who are bound by confidentiality obligations no less protective than those in this Agreement; (c) use such Confidential Information solely for the purposes of performing its obligations or exercising its rights under this Agreement; and (d) protect such Confidential Information using the same degree of care it uses to protect its own confidential information of a similar nature, but in no event less than reasonable care.')

    add_body(doc, '9.2 Exceptions. The obligations of Section 9.1 shall not apply to information that: (a) is or becomes publicly available through no fault of the Receiving Party; (b) was known to the Receiving Party prior to disclosure by the Disclosing Party, as evidenced by written records; (c) is independently developed by the Receiving Party without reference to or use of the Confidential Information; or (d) is rightfully received from a third party without restriction on disclosure.')

    add_body(doc, '9.3 Compelled Disclosure. If the Receiving Party is compelled by law, regulation, or legal process to disclose Confidential Information, the Receiving Party shall: (a) provide the Disclosing Party with prompt written notice to the extent legally permitted; (b) cooperate with the Disclosing Party in seeking a protective order or other appropriate remedy; and (c) disclose only such portion of the Confidential Information as is legally required.')

    add_body(doc, '9.4 Duration. The confidentiality obligations under this Article 9 shall survive the expiration or termination of this Agreement for a period of five (5) years; provided, however, that obligations regarding trade secrets shall continue for so long as the information qualifies as a trade secret under applicable law.')

    # ---- ARTICLE 10: DATA PROTECTION ----
    add_heading_para(doc, 'ARTICLE 10: DATA PROTECTION AND SECURITY', level=1)

    add_body(doc, '10.1 Compliance. Each party shall comply with all applicable data protection and privacy laws and regulations, including without limitation the California Consumer Privacy Act (CCPA), the General Data Protection Regulation (GDPR) to the extent applicable, and any other applicable federal, state, or international data protection laws.')

    add_body(doc, '10.2 Data Processing. To the extent Acme processes personal data on behalf of Beta, the parties shall execute a Data Processing Agreement substantially in the form attached as Exhibit C. Acme shall process personal data only as necessary to perform the Services and in accordance with Beta\'s documented instructions.')

    add_body(doc, '10.3 Security Measures. Acme shall implement and maintain appropriate technical and organizational security measures to protect Beta\'s data against unauthorized access, disclosure, alteration, or destruction. Such measures shall include, at a minimum: (a) encryption of data in transit and at rest; (b) access controls and authentication mechanisms; (c) regular security assessments and penetration testing; (d) incident response procedures; and (e) employee security awareness training.')

    add_body(doc, '10.4 Data Breach Notification. Acme shall notify Beta of any actual or suspected security breach involving Beta\'s data within twenty-four (24) hours of discovery. Such notice shall include a description of the nature of the breach, the categories and approximate number of data subjects affected, the likely consequences of the breach, and the measures taken or proposed to be taken to address the breach and mitigate its effects.')

    # ---- ARTICLE 11: REPRESENTATIONS AND WARRANTIES ----
    add_heading_para(doc, 'ARTICLE 11: REPRESENTATIONS AND WARRANTIES', level=1)

    add_body(doc, '11.1 Mutual Representations. Each party represents and warrants that: (a) it is duly organized, validly existing, and in good standing under the laws of its jurisdiction of organization; (b) it has full corporate power and authority to enter into this Agreement and perform its obligations hereunder; (c) this Agreement constitutes a legal, valid, and binding obligation enforceable against it in accordance with its terms; and (d) the execution, delivery, and performance of this Agreement does not conflict with any other agreement or obligation to which it is a party.')

    add_body(doc, '11.2 Acme Warranties. Acme further represents and warrants that: (a) the Services will be performed in a professional and workmanlike manner consistent with generally accepted industry standards; (b) Acme personnel possess the qualifications, skills, and experience necessary to perform the assigned Services; (c) the Deliverables will conform to the specifications and acceptance criteria set forth in the applicable SOW; (d) the Deliverables will not infringe or misappropriate any third party\'s Intellectual Property rights; and (e) Acme will comply with all applicable laws, regulations, and industry standards in performing the Services.')

    add_body(doc, '11.3 Disclaimer. EXCEPT AS EXPRESSLY SET FORTH IN THIS AGREEMENT, NEITHER PARTY MAKES ANY WARRANTIES, EXPRESS OR IMPLIED, INCLUDING WITHOUT LIMITATION ANY IMPLIED WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE, OR NON-INFRINGEMENT.')

    # ---- ARTICLE 12: WARRANTY PERIOD ----
    add_heading_para(doc, 'ARTICLE 12: WARRANTY AND MAINTENANCE', level=1)

    add_body(doc, '12.1 Warranty Period. Acme warrants that each Deliverable will conform to its specifications for a period of ninety (90) days following acceptance (the "Warranty Period"). During the Warranty Period, Acme shall correct, at no additional charge, any material defects or nonconformities reported by Beta in writing.')

    add_body(doc, '12.2 Maintenance Support. Following the expiration of the Warranty Period, Beta may engage Acme to provide ongoing maintenance and support services under a separate SOW or support agreement. Maintenance services may include bug fixes, patches, updates, and technical assistance as specified therein.')

    # ---- ARTICLE 13: INDEMNIFICATION ----
    add_heading_para(doc, 'ARTICLE 13: INDEMNIFICATION', level=1)

    add_body(doc, '13.1 By Acme. Acme shall indemnify, defend, and hold harmless Beta and its officers, directors, employees, and agents from and against any and all third-party claims, damages, losses, liabilities, costs, and expenses (including reasonable attorneys\' fees) arising out of or relating to: (a) Acme\'s breach of any representation, warranty, or obligation under this Agreement; (b) any claim that the Deliverables or Services infringe or misappropriate a third party\'s Intellectual Property rights; (c) Acme\'s negligence or willful misconduct in the performance of Services; or (d) any violation of applicable law by Acme or its personnel in connection with this Agreement.')

    add_body(doc, '13.2 By Beta. Beta shall indemnify, defend, and hold harmless Acme and its officers, directors, employees, and agents from and against any and all third-party claims, damages, losses, liabilities, costs, and expenses (including reasonable attorneys\' fees) arising out of or relating to: (a) Beta\'s breach of any representation, warranty, or obligation under this Agreement; (b) Beta\'s negligence or willful misconduct; or (c) any materials or information provided by Beta to Acme that infringe or misappropriate a third party\'s Intellectual Property rights.')

    add_body(doc, '13.3 Indemnification Procedure. The indemnified party shall: (a) give the indemnifying party prompt written notice of any claim; (b) grant the indemnifying party sole control of the defense and settlement of the claim; and (c) provide reasonable cooperation and assistance at the indemnifying party\'s expense. The indemnifying party shall not settle any claim in a manner that imposes obligations on the indemnified party or admits liability on behalf of the indemnified party without the indemnified party\'s prior written consent.')

    # ---- ARTICLE 14: TERMINATION ----
    add_heading_para(doc, 'ARTICLE 14: TERMINATION', level=1)

    add_body(doc, '14.1 Termination for Convenience. Either party may terminate this Agreement or any SOW for convenience upon sixty (60) days\' prior written notice to the other party. In the event of termination for convenience, Beta shall pay Acme for all Services rendered and approved expenses incurred through the effective date of termination, plus any non-cancellable commitments made by Acme in reliance on the SOW.')

    add_body(doc, '14.2 Termination for Cause. Either party may terminate this Agreement or any SOW immediately upon written notice if the other party: (a) materially breaches this Agreement and fails to cure such breach within thirty (30) days after receiving written notice thereof; (b) becomes insolvent, files for bankruptcy, or has a receiver appointed for a substantial part of its assets; or (c) ceases to conduct business in the ordinary course.')

    add_body(doc, '14.3 Effect of Termination. Upon expiration or termination of this Agreement for any reason: (a) each party shall promptly return or destroy all Confidential Information of the other party; (b) Acme shall deliver to Beta all completed or in-progress Deliverables and Work Product; (c) Beta shall pay all undisputed amounts owed to Acme as of the termination date; and (d) the following provisions shall survive: Articles 8, 9, 10, 11.3, 13, 15, 16, 17, and 18.')

    # ---- ARTICLE 15: LIMITATION OF LIABILITY ----
    add_heading_para(doc, 'ARTICLE 15: LIMITATION OF LIABILITY', level=1)

    add_body(doc, '15.1 Consequential Damages. IN NO EVENT SHALL EITHER PARTY BE LIABLE TO THE OTHER FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, PUNITIVE, OR EXEMPLARY DAMAGES, INCLUDING BUT NOT LIMITED TO DAMAGES FOR LOSS OF PROFITS, GOODWILL, DATA, OR BUSINESS OPPORTUNITIES, REGARDLESS OF WHETHER SUCH DAMAGES ARE BASED ON WARRANTY, CONTRACT, TORT, STRICT LIABILITY, OR ANY OTHER LEGAL THEORY, AND EVEN IF THE PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.')

    add_body(doc, '15.2 Cap on Liability. EXCEPT FOR OBLIGATIONS UNDER ARTICLES 9 (CONFIDENTIALITY) AND 13 (INDEMNIFICATION), EACH PARTY\'S TOTAL AGGREGATE LIABILITY UNDER THIS AGREEMENT SHALL NOT EXCEED THE TOTAL FEES PAID OR PAYABLE BY BETA TO ACME UNDER THIS AGREEMENT DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE DATE OF THE CLAIM GIVING RISE TO SUCH LIABILITY.')

    add_body(doc, '15.3 Exceptions. The limitations set forth in Sections 15.1 and 15.2 shall not apply to: (a) liability arising from a party\'s willful misconduct or gross negligence; (b) liability arising from a breach of Article 9 (Confidentiality); (c) Acme\'s indemnification obligations under Section 13.1(b) (IP infringement); or (d) liability that cannot be limited under applicable law.')

    # ---- ARTICLE 16: DISPUTE RESOLUTION ----
    add_heading_para(doc, 'ARTICLE 16: DISPUTE RESOLUTION', level=1)

    add_body(doc, '16.1 Negotiation. The parties shall attempt to resolve any dispute arising out of or relating to this Agreement through good faith negotiation between senior executives of each party. Either party may initiate such negotiation by providing written notice to the other party describing the dispute and proposing a resolution. The executives shall meet (in person or by videoconference) within fifteen (15) business days of receipt of such notice.')

    add_body(doc, '16.2 Mediation. If the dispute cannot be resolved through negotiation within thirty (30) days, either party may submit the dispute to non-binding mediation administered by the American Arbitration Association under its Commercial Mediation Procedures. The mediation shall be conducted in San Francisco, California, and each party shall bear its own costs of mediation, with the mediator\'s fees shared equally.')

    add_body(doc, '16.3 Arbitration. If the dispute cannot be resolved through mediation within sixty (60) days of the commencement of mediation, either party may submit the dispute to binding arbitration administered by the American Arbitration Association under its Commercial Arbitration Rules. The arbitration shall be conducted by a panel of three (3) arbitrators in San Francisco, California. The arbitrators\' decision shall be final and binding, and judgment on the award may be entered in any court of competent jurisdiction.')

    add_body(doc, '16.4 Injunctive Relief. Notwithstanding the foregoing, either party may seek injunctive or other equitable relief in any court of competent jurisdiction to prevent irreparable harm pending the resolution of any dispute under this Article 16, including without limitation enforcement of the confidentiality obligations under Article 9 or protection of Intellectual Property rights under Article 8.')

    # ---- ARTICLE 17: INSURANCE ----
    add_heading_para(doc, 'ARTICLE 17: INSURANCE', level=1)

    add_body(doc, '17.1 Required Coverage. During the Term and for a period of two (2) years thereafter, Acme shall maintain the following insurance coverage: (a) Commercial General Liability with minimum limits of $2,000,000 per occurrence and $5,000,000 in the aggregate; (b) Professional Liability (Errors and Omissions) with minimum limits of $5,000,000 per claim and in the aggregate; (c) Workers\' Compensation as required by applicable law; (d) Employer\'s Liability with minimum limits of $1,000,000 per occurrence; and (e) Cyber Liability with minimum limits of $5,000,000 per claim and in the aggregate.')

    add_body(doc, '17.2 Certificates. Upon Beta\'s request, Acme shall provide certificates of insurance evidencing the required coverage, naming Beta as an additional insured on the Commercial General Liability and Cyber Liability policies. Acme shall provide at least thirty (30) days\' prior written notice to Beta of any material change in, cancellation of, or failure to renew any required insurance coverage.')

    # ---- ARTICLE 18: GENERAL PROVISIONS ----
    add_heading_para(doc, 'ARTICLE 18: GENERAL PROVISIONS', level=1)

    add_body(doc, '18.1 Governing Law. This Agreement shall be governed by and construed in accordance with the laws of the State of California, without regard to its conflicts of laws principles.')

    add_body(doc, '18.2 Entire Agreement. This Agreement, including all Exhibits, Schedules, and Statements of Work, constitutes the entire agreement between the parties with respect to the subject matter hereof and supersedes all prior and contemporaneous agreements, proposals, representations, and understandings, whether written or oral.')

    add_body(doc, '18.3 Notices. All notices required or permitted under this Agreement shall be in writing and shall be deemed given when: (a) delivered personally; (b) sent by confirmed overnight courier; (c) sent by certified mail, return receipt requested; or (d) sent by email with confirmed receipt. Notices shall be addressed to the parties at the addresses set forth in Exhibit D or at such other address as either party may designate by written notice.')

    add_body(doc, '18.4 Assignment. Neither party may assign its rights or delegate its obligations under this Agreement without the prior written consent of the other party, except that either party may assign this Agreement to an Affiliate or in connection with a merger, acquisition, or sale of all or substantially all of its assets, provided that the assignee agrees in writing to be bound by the terms of this Agreement.')

    add_body(doc, '18.5 Amendments. This Agreement may not be modified or amended except by a written instrument signed by authorized representatives of both parties. No waiver of any provision shall be effective unless made in writing and signed by the waiving party.')

    add_body(doc, '18.6 Severability. If any provision of this Agreement is held to be invalid, illegal, or unenforceable, the remaining provisions shall continue in full force and effect. The parties shall negotiate in good faith to replace the invalid provision with a valid provision that most closely approximates the intent and economic effect of the invalid provision.')

    add_body(doc, '18.7 Force Majeure. Neither party shall be liable for any delay or failure to perform its obligations (other than payment obligations) due to a Force Majeure Event. The affected party shall notify the other party promptly and shall use commercially reasonable efforts to resume performance as soon as practicable. If a Force Majeure Event continues for more than ninety (90) days, either party may terminate the affected SOW upon written notice.')

    add_body(doc, '18.8 Independent Contractor. Acme is an independent contractor and nothing in this Agreement shall be construed to create a partnership, joint venture, agency, or employment relationship between the parties. Acme shall be solely responsible for all taxes, withholdings, and other statutory or contractual obligations of any sort, including workers\' compensation insurance premiums and unemployment compensation contributions.')

    add_body(doc, '18.9 Waiver. The failure of either party to enforce any provision of this Agreement shall not be deemed a waiver of future enforcement of that or any other provision. All waivers must be in writing and signed by the waiving party to be effective.')

    add_body(doc, '18.10 Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an original, and all of which together shall constitute one and the same instrument. Electronic signatures shall be deemed original signatures for all purposes.')

    # ---- SIGNATURE PAGE ----
    doc.add_paragraph()
    doc.add_paragraph()

    add_body(doc, 'IN WITNESS WHEREOF, the parties hereto have caused this Master Services Agreement to be executed by their duly authorized representatives as of the Effective Date.', bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    # Acme signature block
    add_body(doc, 'ACME CORPORATION', bold=True)
    doc.add_paragraph()
    add_body(doc, '________________________________________')
    add_body(doc, 'Name: Victoria R. Harrington')
    add_body(doc, 'Title: Chief Executive Officer')
    add_body(doc, 'Date: January 15, 2025')

    doc.add_paragraph()
    doc.add_paragraph()

    # Beta signature block
    add_body(doc, 'BETA INCORPORATED', bold=True)
    doc.add_paragraph()
    add_body(doc, '________________________________________')
    add_body(doc, 'Name: Jonathan M. Whitfield')
    add_body(doc, 'Title: Chief Operating Officer')
    add_body(doc, 'Date: January 15, 2025')

    doc.add_paragraph()

    # ---- EXHIBITS ----
    doc.add_page_break()

    add_heading_para(doc, 'EXHIBIT A: STATEMENT OF WORK TEMPLATE', level=1)

    add_body(doc, 'Each Statement of Work shall include the following sections:')
    doc.add_paragraph('1. SOW Reference Number and Effective Date', style='List Number')
    doc.add_paragraph('2. Project Description and Objectives', style='List Number')
    doc.add_paragraph('3. Scope of Services', style='List Number')
    doc.add_paragraph('4. Deliverables and Acceptance Criteria', style='List Number')
    doc.add_paragraph('5. Project Timeline and Milestones', style='List Number')
    doc.add_paragraph('6. Resource Allocation and Key Personnel', style='List Number')
    doc.add_paragraph('7. Fees and Payment Schedule', style='List Number')
    doc.add_paragraph('8. Assumptions and Dependencies', style='List Number')
    doc.add_paragraph('9. Risk Factors and Mitigation Strategies', style='List Number')
    doc.add_paragraph('10. Special Terms and Conditions', style='List Number')

    doc.add_page_break()

    add_heading_para(doc, 'EXHIBIT B: RATE CARD', level=1)

    # Rate card table
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    headers = ['Role', 'Hourly Rate (USD)', 'Daily Rate (USD)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    rates = [
        ('Senior Architect', '$325', '$2,600'),
        ('Technical Lead', '$275', '$2,200'),
        ('Senior Developer', '$225', '$1,800'),
        ('Developer', '$175', '$1,400'),
        ('QA Engineer', '$150', '$1,200'),
        ('Project Manager', '$250', '$2,000'),
        ('Business Analyst', '$200', '$1,600'),
    ]
    for r, (role, hourly, daily) in enumerate(rates, 1):
        table.cell(r, 0).text = role
        table.cell(r, 1).text = hourly
        table.cell(r, 2).text = daily

    add_body(doc, '')
    add_body(doc, 'Rates are effective as of the Effective Date and subject to annual adjustment per Section 7.6.')

    doc.add_page_break()

    add_heading_para(doc, 'EXHIBIT C: DATA PROCESSING AGREEMENT', level=1)

    add_body(doc, 'This Data Processing Agreement ("DPA") supplements the Master Services Agreement between Acme Corporation and Beta Incorporated.')

    add_body(doc, '1. Definitions. Capitalized terms not defined herein shall have the meanings assigned to them in the Agreement. "Personal Data" means any information relating to an identified or identifiable natural person, as defined by applicable data protection laws.')

    add_body(doc, '2. Processing Purpose. Acme shall process Personal Data solely for the purpose of providing the Services described in the Agreement and applicable SOWs, and in accordance with Beta\'s documented instructions.')

    add_body(doc, '3. Sub-processing. Acme shall not engage any sub-processor to process Personal Data without Beta\'s prior written consent. Acme shall maintain a list of approved sub-processors and shall notify Beta of any changes thereto at least thirty (30) days in advance.')

    add_body(doc, '4. Data Subject Rights. Acme shall assist Beta in fulfilling its obligations to respond to data subject requests, including requests for access, rectification, erasure, restriction, portability, and objection, within the timeframes required by applicable law.')

    add_body(doc, '5. Security. Acme shall implement the security measures described in Section 10.3 of the Agreement and shall conduct annual audits of its security practices. Acme shall provide Beta with audit reports upon request.')

    add_body(doc, '6. International Transfers. Acme shall not transfer Personal Data to any country outside the United States without Beta\'s prior written consent and without ensuring appropriate safeguards are in place, such as Standard Contractual Clauses or binding corporate rules approved by the relevant supervisory authority.')

    doc.add_page_break()

    add_heading_para(doc, 'EXHIBIT D: NOTICE ADDRESSES', level=1)

    add_body(doc, 'Acme Corporation', bold=True)
    add_body(doc, 'Attn: Legal Department')
    add_body(doc, '1250 Innovation Drive, Suite 400')
    add_body(doc, 'Palo Alto, California 94301')
    add_body(doc, 'Email: legal@acmecorp.com')

    doc.add_paragraph()

    add_body(doc, 'Beta Incorporated', bold=True)
    add_body(doc, 'Attn: General Counsel')
    add_body(doc, '800 Market Street, 12th Floor')
    add_body(doc, 'San Francisco, California 94102')
    add_body(doc, 'Email: legal@betainc.com')

    doc.add_paragraph()
    add_body(doc, '* * * End of Master Services Agreement * * *', bold=True)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
