"""
Reward Script: Add running headers to MSA contract with different even/odd/first page headers.
Task ID: writer_legal_057
Domain: libreoffice_writer
Scoring:
  C1: evenAndOddHeaders enabled in document settings (0.15)
  C2: different_first_page_header_footer enabled (0.15)
  C3: First page header is empty / no header on page 1 (0.15)
  C4: Even page header text = 'Master Services Agreement' (0.20)
  C5: Even page header alignment = LEFT (0.10)
  C6: Odd page (default) header text = 'Acme Corp / Beta Inc.' (0.15)
  C7: Odd page (default) header alignment = RIGHT (0.10)
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_057'


def persist_app_state():
    """Best-effort save via Ctrl+S in case doc is open in LibreOffice."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify header configuration with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError as e:
        print(f"CRITICAL: Missing library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if len(doc.sections) == 0:
        print("CRITICAL: Document has no sections")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: evenAndOddHeaders enabled in document settings (0.15 points)
    try:
        settings_elem = doc.settings.element
        eaoh = settings_elem.find(qn('w:evenAndOddHeaders'))
        if eaoh is not None:
            # Element present means enabled (val=None or val="true" both mean enabled)
            val = eaoh.get(qn('w:val'))
            if val is None or val.lower() in ('true', '1', 'on'):
                print(f"PASS: Component 1 — evenAndOddHeaders is enabled (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 1 — evenAndOddHeaders element present but val={val}")
        else:
            print("FAIL: Component 1 — evenAndOddHeaders element not found in settings")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: different_first_page_header_footer enabled (0.15 points)
    try:
        if section.different_first_page_header_footer:
            print(f"PASS: Component 2 — different_first_page_header_footer is True (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 — different_first_page_header_footer is False")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: different_first_page enabled AND first page header is empty (0.15 points)
    # Both sub-conditions must hold: the feature must be turned on AND the first page header must be blank.
    # This ensures initial_env (where different_first_page is False) does not score.
    try:
        first_header = section.first_page_header
        first_header_text = ""
        for p in first_header.paragraphs:
            first_header_text += p.text.strip()
        if section.different_first_page_header_footer and first_header_text == "":
            print(f"PASS: Component 3 — different_first_page enabled and first page header is empty (0.15 pts)")
            total_score += 0.15
        elif not section.different_first_page_header_footer:
            print(f"FAIL: Component 3 — different_first_page_header_footer is not enabled")
        else:
            print(f"FAIL: Component 3 — First page header has text: {repr(first_header_text)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Even page header text = 'Master Services Agreement' (0.20 points)
    try:
        even_header = section.even_page_header
        even_text = ""
        for p in even_header.paragraphs:
            t = p.text.strip()
            if t:
                even_text = t
                break
        if even_text == "Master Services Agreement":
            print(f"PASS: Component 4 — Even page header text matches exactly (0.20 pts)")
            total_score += 0.20
        elif "master services agreement" in even_text.lower():
            print(f"PARTIAL: Component 4 — Even page header text case mismatch: {repr(even_text)} (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — Even page header text: {repr(even_text)}, expected 'Master Services Agreement'")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Even page header alignment = LEFT (0.10 points)
    try:
        even_header = section.even_page_header
        even_align = None
        for p in even_header.paragraphs:
            if p.text.strip():
                even_align = p.paragraph_format.alignment
                break
        if even_align == WD_PARAGRAPH_ALIGNMENT.LEFT or even_align is None:
            # None means default which is LEFT
            # But we need to be careful: check if text is present first
            even_text_check = "".join(p.text.strip() for p in even_header.paragraphs)
            if even_text_check:
                print(f"PASS: Component 5 — Even page header is left-aligned (alignment={even_align}) (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 5 — Even page header has no text, alignment check not meaningful")
        else:
            print(f"FAIL: Component 5 — Even page header alignment: {even_align}, expected LEFT")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Odd page (default) header text = 'Acme Corp / Beta Inc.' (0.15 points)
    try:
        default_header = section.header
        odd_text = ""
        for p in default_header.paragraphs:
            t = p.text.strip()
            if t:
                odd_text = t
                break
        if odd_text == "Acme Corp / Beta Inc.":
            print(f"PASS: Component 6 — Odd page header text matches exactly (0.15 pts)")
            total_score += 0.15
        elif "acme" in odd_text.lower() and "beta" in odd_text.lower():
            print(f"PARTIAL: Component 6 — Odd page header text close match: {repr(odd_text)} (0.07 pts)")
            total_score += 0.07
        else:
            print(f"FAIL: Component 6 — Odd page (default) header text: {repr(odd_text)}, expected 'Acme Corp / Beta Inc.'")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Odd page (default) header alignment = RIGHT (0.10 points)
    try:
        default_header = section.header
        odd_align = None
        for p in default_header.paragraphs:
            if p.text.strip():
                odd_align = p.paragraph_format.alignment
                break
        if odd_align == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            print(f"PASS: Component 7 — Odd page header is right-aligned (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 7 — Odd page header alignment: {odd_align}, expected RIGHT (2)")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
