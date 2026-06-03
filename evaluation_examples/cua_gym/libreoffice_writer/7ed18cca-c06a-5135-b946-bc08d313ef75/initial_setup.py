"""
Initial Setup: Configuration Options table with plain-text Allowed Values
Task ID: writer_tech_070
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_070'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    heading = doc.add_heading('Application Configuration Reference', level=1)

    # Intro paragraph
    doc.add_paragraph(
        'This document describes the configuration options available for the '
        'DataStream Analytics Platform v3.2. Each option controls a specific '
        'aspect of the application behavior. Refer to the Allowed Values column '
        'for valid settings.'
    )

    doc.add_paragraph()  # spacer

    # --- Configuration Options Table ---
    # Columns: Option Name | Type | Default Value | Allowed Values | Description
    config_data = [
        [
            'log_level',
            'enum',
            'INFO',
            'DEBUG, INFO, WARNING, ERROR, CRITICAL',
            'Sets the verbosity of application logging output.'
        ],
        [
            'output_format',
            'enum',
            'JSON',
            'JSON, XML, CSV, YAML, PARQUET',
            'Determines the file format for exported analytics reports.'
        ],
        [
            'auth_mode',
            'enum',
            'TOKEN',
            'TOKEN, OAUTH2, SAML, LDAP, API_KEY',
            'Specifies the authentication mechanism for API access.'
        ],
        [
            'cache_strategy',
            'enum',
            'LRU',
            'LRU, LFU, FIFO, TTL, NONE',
            'Controls the eviction policy used by the in-memory cache.'
        ],
        [
            'compression_type',
            'enum',
            'GZIP',
            'GZIP, ZSTD, LZ4, SNAPPY, NONE',
            'Selects the compression algorithm for data storage.'
        ],
        [
            'retry_policy',
            'enum',
            'EXPONENTIAL',
            'EXPONENTIAL, LINEAR, FIXED, NONE',
            'Defines the backoff strategy for failed network requests.'
        ],
        [
            'db_isolation_level',
            'enum',
            'READ_COMMITTED',
            'READ_UNCOMMITTED, READ_COMMITTED, REPEATABLE_READ, SERIALIZABLE',
            'Sets the transaction isolation level for database operations.'
        ],
        [
            'thread_pool_mode',
            'enum',
            'DYNAMIC',
            'DYNAMIC, FIXED, CACHED, SINGLE',
            'Configures how the worker thread pool is managed.'
        ],
    ]

    headers = ['Option Name', 'Type', 'Default Value', 'Allowed Values', 'Description']
    table = doc.add_table(rows=1 + len(config_data), cols=5)
    table.style = 'Table Grid'

    # Header row formatting
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Dark blue background via shading
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '2F5496',
        })
        shading.append(shd)

    # Data rows
    for row_idx, row_data in enumerate(config_data, start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(9)
            # Make option name monospace
            if col_idx == 0:
                run.font.name = 'Courier New'
            # Make type column italic
            if col_idx == 1:
                run.italic = True

    # Set approximate column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.3)
        row.cells[1].width = Inches(0.6)
        row.cells[2].width = Inches(1.1)
        row.cells[3].width = Inches(2.2)
        row.cells[4].width = Inches(2.3)

    # --- Notes section ---
    doc.add_paragraph()
    doc.add_heading('Notes', level=2)
    doc.add_paragraph(
        'All enum values are case-sensitive. Using an invalid value will cause '
        'the application to fall back to the default setting and emit a warning '
        'in the log output.'
    )
    doc.add_paragraph(
        'Configuration changes require a service restart unless hot-reload is '
        'enabled in the deployment manifest.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
