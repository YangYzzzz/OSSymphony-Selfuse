"""
Reward Script: Create nested table inside 'Configuration Options' table
Task ID: writer_tech_070
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): At least one 'Allowed Values' cell contains a nested table
  Component 2 (0.3): First nested table has correct 2-column structure with Value/Description headers
                      and rows matching the original enum values
  Component 3 (0.2): Nested table descriptions are non-empty for each enum value
  Component 4 (0.2): At least two rows have nested tables (both log_level and output_format converted)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_070'

# Expected enum values from the initial state (comma-separated in Allowed Values column)
EXPECTED_ENUMS = {
    'log_level': ['DEBUG', 'INFO', 'WARNING', 'ERROR', 'CRITICAL'],
    'output_format': ['JSON', 'XML', 'CSV', 'YAML', 'PARQUET'],
    'auth_mode': ['TOKEN', 'OAUTH2', 'SAML', 'LDAP', 'API_KEY'],
    'cache_strategy': ['LRU', 'LFU', 'FIFO', 'TTL', 'NONE'],
    'compression_type': ['GZIP', 'ZSTD', 'LZ4', 'SNAPPY', 'NONE'],
    'retry_policy': ['EXPONENTIAL', 'LINEAR', 'FIXED', 'NONE'],
    'db_isolation_level': ['READ_UNCOMMITTED', 'READ_COMMITTED', 'REPEATABLE_READ', 'SERIALIZABLE'],
    'thread_pool_mode': ['DYNAMIC', 'FIXED', 'CACHED', 'SINGLE'],
}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table (the Configuration Options table)
    if len(doc.tables) < 1:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    main_table = doc.tables[0]

    # Precondition: table must have the expected header structure
    try:
        header_row = main_table.rows[0]
        headers = [cell.text.strip() for cell in header_row.cells]
        if 'Allowed Values' not in headers:
            print(f"FAIL: 'Allowed Values' column not found. Headers: {headers}")
            print("REWARD: 0.0")
            return 0.0
        allowed_col_idx = headers.index('Allowed Values')
        option_col_idx = headers.index('Option Name') if 'Option Name' in headers else 0
    except Exception as e:
        print(f"ERROR: Cannot parse table headers: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find all data rows that have nested tables in the Allowed Values column
    rows_with_nested = []  # list of (row_index, option_name, nested_table)
    for ri in range(1, len(main_table.rows)):
        try:
            cell = main_table.cell(ri, allowed_col_idx)
            option_name = main_table.cell(ri, option_col_idx).text.strip()
            nested_tables = cell.tables
            if len(nested_tables) > 0:
                rows_with_nested.append((ri, option_name, nested_tables[0]))
        except Exception as e:
            print(f"ERROR: Cannot check row {ri}: {e}")

    print(f"Found {len(rows_with_nested)} row(s) with nested tables in 'Allowed Values' column")

    # Component 1: At least one 'Allowed Values' cell contains a nested table (0.3 points)
    # This is the core task requirement - converting plain text to a nested table
    try:
        if len(rows_with_nested) >= 1:
            names = [r[1] for r in rows_with_nested]
            print(f"PASS: Component 1 - Found nested table(s) in rows: {names} (0.3 pts)")
            total_score += 0.3
        else:
            print("FAIL: Component 1 - No nested tables found in any 'Allowed Values' cell")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: First nested table has correct 2-column structure with Value/Description
    # headers and rows matching the expected enum values (0.3 points)
    try:
        if len(rows_with_nested) >= 1:
            _, option_name, nested_tbl = rows_with_nested[0]
            nt_rows = len(nested_tbl.rows)
            nt_cols = len(nested_tbl.columns)

            # Check column count is 2
            col_ok = (nt_cols == 2)

            # Check headers contain "Value" and "Description" (case-insensitive)
            header_cells = [nested_tbl.cell(0, ci).text.strip().lower() for ci in range(nt_cols)]
            headers_ok = ('value' in header_cells and 'description' in header_cells)

            # Check that the enum values from the original comma-separated list are present
            expected_vals = EXPECTED_ENUMS.get(option_name, [])
            found_vals = []
            for nri in range(1, nt_rows):
                val_text = nested_tbl.cell(nri, 0).text.strip()
                found_vals.append(val_text)

            # Check if all expected enum values appear in the nested table
            vals_match = all(ev in found_vals for ev in expected_vals) if expected_vals else False

            if col_ok and headers_ok and vals_match:
                print(f"PASS: Component 2 - Nested table for '{option_name}' has correct structure: "
                      f"2 cols, Value/Description headers, {len(found_vals)} enum values match (0.3 pts)")
                total_score += 0.3
            else:
                details = []
                if not col_ok:
                    details.append(f"expected 2 cols, got {nt_cols}")
                if not headers_ok:
                    details.append(f"headers {header_cells} missing Value/Description")
                if not vals_match:
                    missing = [v for v in expected_vals if v not in found_vals]
                    details.append(f"missing enum values: {missing}, found: {found_vals}")
                print(f"FAIL: Component 2 - Nested table structure issues: {'; '.join(details)}")
        else:
            print("FAIL: Component 2 - No nested tables found (depends on Component 1)")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Nested table descriptions are non-empty for each enum value (0.2 points)
    # Each row in the nested table should have a meaningful description
    try:
        if len(rows_with_nested) >= 1:
            _, option_name, nested_tbl = rows_with_nested[0]
            nt_rows = len(nested_tbl.rows)
            desc_col = 1  # Description is second column

            non_empty_descs = 0
            total_data_rows = nt_rows - 1  # minus header

            for nri in range(1, nt_rows):
                desc_text = nested_tbl.cell(nri, desc_col).text.strip()
                if len(desc_text) > 3:  # meaningful description, not just a letter
                    non_empty_descs += 1

            if total_data_rows > 0 and non_empty_descs == total_data_rows:
                print(f"PASS: Component 3 - All {non_empty_descs}/{total_data_rows} enum values "
                      f"have descriptions (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 - Only {non_empty_descs}/{total_data_rows} enum values "
                      f"have descriptions (need all)")
        else:
            print("FAIL: Component 3 - No nested tables found (depends on Component 1)")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: At least two rows have nested tables with proper structure (0.2 points)
    # The golden state has both log_level and output_format converted
    try:
        valid_nested_count = 0
        for _, option_name, nested_tbl in rows_with_nested:
            nt_cols = len(nested_tbl.columns)
            nt_rows = len(nested_tbl.rows)
            if nt_cols == 2 and nt_rows >= 2:
                # Has at least a header row and one data row with 2 columns
                header_cells = [nested_tbl.cell(0, ci).text.strip().lower() for ci in range(nt_cols)]
                if 'value' in header_cells and 'description' in header_cells:
                    valid_nested_count += 1

        if valid_nested_count >= 2:
            print(f"PASS: Component 4 - {valid_nested_count} rows have properly structured "
                  f"nested tables (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 4 - Only {valid_nested_count} row(s) have valid nested tables, "
                  f"need at least 2")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
