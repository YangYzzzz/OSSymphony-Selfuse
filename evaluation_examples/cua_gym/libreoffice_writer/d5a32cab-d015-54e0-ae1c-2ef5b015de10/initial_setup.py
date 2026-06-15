"""
Initial Setup: Attendance table document with 4 rows (header + 3 data rows)
Task ID: writer_tbl_016
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'  # VM path — file on Desktop
TASK_ID = 'attendance'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Add a brief title paragraph
    heading = doc.add_paragraph("Attendance Log")
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(14)

    # Create the attendance table: 4 rows x 4 columns
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    # Row 1 (header)
    header_row = table.rows[0]
    header_row.cells[0].text = 'Week'
    header_row.cells[1].text = 'Topic'
    header_row.cells[2].text = 'Date'
    header_row.cells[3].text = 'Attendees'

    # Make header bold
    for cell in header_row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Row 2: Week 1
    row2 = table.rows[1]
    row2.cells[0].text = 'Week 1'
    row2.cells[1].text = 'Introduction'
    row2.cells[2].text = '2024-01-15'
    row2.cells[3].text = '42'

    # Row 3: Week 2
    row3 = table.rows[2]
    row3.cells[0].text = 'Week 2'
    row3.cells[1].text = 'Basics'
    row3.cells[2].text = '2024-01-22'
    row3.cells[3].text = '40'

    # Row 4: Week 3
    row4 = table.rows[3]
    row4.cells[0].text = 'Week 3'
    row4.cells[1].text = 'Advanced'
    row4.cells[2].text = '2024-01-29'
    row4.cells[3].text = '38'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
