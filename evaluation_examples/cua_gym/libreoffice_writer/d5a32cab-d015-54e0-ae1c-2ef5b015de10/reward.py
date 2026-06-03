"""
Reward Script: Add a new row above the first data row in the attendance table
Task ID: writer_tbl_016
Domain: libreoffice_writer
Scoring:
  Component 1: Table has 5 rows (row count increased from 4 to 5)         — 0.3 pts
  Component 2: New row 2 contains correct data (Week 0/Orientation/...)   — 0.5 pts
  Component 3: Original data rows 1-3 still intact in rows 3-5            — 0.2 pts
  Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_016'
FILE_PATH = '/home/user/Desktop/attendance.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.

    Task: Add a new row above the first data row (below the header) in
    the attendance table. Fill it with: 'Week 0', 'Orientation', '2024-01-08', '45'.

    Expected golden state (5 rows):
      Row 0 (header): Week | Topic | Date | Attendees
      Row 1 (new):    Week 0 | Orientation | 2024-01-08 | 45
      Row 2:          Week 1 | Introduction | 2024-01-15 | 42
      Row 3:          Week 2 | Basics | 2024-01-22 | 40
      Row 4:          Week 3 | Advanced | 2024-01-29 | 38

    Initial state (4 rows) is used to confirm baseline:
      Row 0 (header): Week | Topic | Date | Attendees
      Row 1:          Week 1 | Introduction | 2024-01-15 | 42
      Row 2:          Week 2 | Basics | 2024-01-22 | 40
      Row 3:          Week 3 | Advanced | 2024-01-29 | 38

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document — precondition gate
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: exactly 1 table present
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)

    # Component 1: Table now has 5 rows (initial has 4 rows) (0.3 pts)
    try:
        if num_rows == 5:
            print(f"PASS: Component 1 — Table has 5 rows (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — Expected 5 rows, found {num_rows}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: New row 2 (index 1) contains correct data (0.5 pts)
    # Expected: 'Week 0' | 'Orientation' | '2024-01-08' | '45'
    try:
        if num_rows >= 2:
            new_row_cells = [table.rows[1].cells[j].text.strip() for j in range(4)]
            expected_new_row = ['Week 0', 'Orientation', '2024-01-08', '45']
            if new_row_cells == expected_new_row:
                print(f"PASS: Component 2 — New row 2 has correct data: {new_row_cells} (0.5 pts)")
                total_score += 0.5
            else:
                print(f"FAIL: Component 2 — Row 2 expected {expected_new_row}, found {new_row_cells}")
        else:
            print(f"FAIL: Component 2 — Table has fewer than 2 rows, cannot check new row")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Original data rows remain intact in rows 3-5 (indices 2-4) (0.2 pts)
    # Expected (original data from initial_env):
    #   Row 2: Week 1 | Introduction | 2024-01-15 | 42
    #   Row 3: Week 2 | Basics | 2024-01-22 | 40
    #   Row 4: Week 3 | Advanced | 2024-01-29 | 38
    try:
        if num_rows >= 5:
            expected_original_rows = [
                ['Week 1', 'Introduction', '2024-01-15', '42'],
                ['Week 2', 'Basics', '2024-01-22', '40'],
                ['Week 3', 'Advanced', '2024-01-29', '38'],
            ]
            actual_original_rows = [
                [table.rows[i].cells[j].text.strip() for j in range(4)]
                for i in range(2, 5)
            ]
            if actual_original_rows == expected_original_rows:
                print(f"PASS: Component 3 — Original data rows intact in rows 3-5 (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — Original rows mismatch.")
                for idx, (expected, actual) in enumerate(zip(expected_original_rows, actual_original_rows)):
                    if expected != actual:
                        print(f"  Row {idx+2}: expected {expected}, found {actual}")
        else:
            print(f"FAIL: Component 3 — Table has fewer than 5 rows, cannot check original rows")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
