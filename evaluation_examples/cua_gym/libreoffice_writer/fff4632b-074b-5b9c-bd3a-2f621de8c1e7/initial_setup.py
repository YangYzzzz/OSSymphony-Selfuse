"""
Initial Setup: Create a 10-page Writer document with Default Page Style on all pages,
headers and footers enabled throughout (including page 1).
Task ID: writer_bs_060
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_060'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# Realistic filler paragraphs for a business strategy document
PARAGRAPHS = [
    ("Executive Summary", 0),
    ("Meridian Technologies has experienced significant growth over the past fiscal year, "
     "with consolidated revenue reaching $287.4 million, representing a 23% increase over "
     "the prior period. This growth was driven primarily by expansion in our cloud services "
     "division and strategic acquisitions in the Asia-Pacific region. Our customer base now "
     "exceeds 4,200 enterprise clients across 38 countries.", None),
    ("The board of directors has approved a revised capital allocation framework that "
     "prioritizes organic growth investments while maintaining our commitment to returning "
     "value to shareholders through quarterly dividends and opportunistic share repurchases.", None),
    ("Market Analysis", 1),
    ("The global enterprise software market continues to evolve rapidly, with total "
     "addressable market projected to reach $1.2 trillion by 2028. Key trends shaping our "
     "industry include the accelerated adoption of artificial intelligence, the shift toward "
     "consumption-based pricing models, and increasing regulatory requirements around data "
     "sovereignty and privacy.", None),
    ("Our competitive landscape analysis identifies three primary segments: established "
     "players like Oracle and SAP who are transitioning legacy portfolios to cloud-native "
     "architectures; mid-market challengers including ServiceNow and Workday who compete on "
     "vertical specialization; and emerging disruptors leveraging open-source foundations to "
     "offer compelling price-performance alternatives.", None),
    ("Regional Performance Overview", 1),
    ("North America remains our largest market, contributing $168.2 million in revenue with "
     "a gross margin of 72%. The region added 340 new enterprise accounts during the period, "
     "with particular strength in financial services and healthcare verticals. Average contract "
     "value increased to $127,000 annually, reflecting successful upselling of premium "
     "analytics and security modules.", None),
    ("The European division generated $74.8 million in revenue, a 19% increase driven by "
     "GDPR compliance solutions and expanding partnerships with regional systems integrators. "
     "Our Frankfurt data center, operational since Q2, has reduced latency for Central European "
     "clients by an average of 40% and enabled compliance with German federal data residency "
     "requirements.", None),
    ("Asia-Pacific revenue reached $44.4 million following the acquisition of CloudBridge "
     "Solutions in Singapore. The integration is proceeding ahead of schedule, with combined "
     "team synergies expected to yield $8 million in annual cost savings by Q3 of the current "
     "fiscal year. Japan and Australia represent our fastest-growing markets in the region.", None),
    ("Product Development Roadmap", 1),
    ("Our engineering organization, now comprising 1,240 professionals across five R&D "
     "centers, delivered 47 major feature releases during the period. Key launches included "
     "the Meridian Intelligence Suite, an AI-powered analytics platform that provides predictive "
     "insights for business operations, and Meridian SecureVault, our zero-trust data protection "
     "solution designed for multi-cloud environments.", None),
    ("The upcoming fiscal year will focus on three strategic pillars: first, deepening our AI "
     "and machine learning capabilities through both organic development and targeted acquisitions; "
     "second, expanding our platform ecosystem with enhanced API frameworks and a revamped developer "
     "portal; and third, launching industry-specific solution accelerators for manufacturing, "
     "retail, and public sector verticals.", None),
    ("Financial Performance Details", 1),
    ("Operating expenses totaled $198.3 million for the period, with research and development "
     "accounting for 28% of revenue. Sales and marketing expenses increased by 15% to support "
     "territory expansion in Latin America and the Middle East. General and administrative costs "
     "were held flat through process automation initiatives that eliminated redundant workflows "
     "across finance, HR, and procurement functions.", None),
    ("Free cash flow generation remained robust at $52.1 million, enabling the completion of "
     "two strategic acquisitions, the repayment of $30 million in revolving credit facility "
     "borrowings, and the distribution of $18.4 million in shareholder dividends. The company "
     "maintains a strong balance sheet with $143 million in cash and short-term investments and "
     "no long-term debt obligations.", None),
    ("Talent and Culture Initiatives", 1),
    ("Total headcount grew to 3,870 employees, a net increase of 520 positions. Employee "
     "engagement scores improved to 4.2 out of 5.0, up from 3.8 in the prior period, reflecting "
     "investments in flexible work arrangements, expanded professional development programs, and "
     "a redesigned compensation framework that increased base salary competitiveness to the 65th "
     "percentile of our peer group.", None),
    ("Our diversity metrics showed meaningful progress, with women representing 38% of "
     "technical roles, up from 31% two years ago. The Meridian Fellows program, which funds "
     "advanced degrees for high-potential employees, expanded to include 45 participants across "
     "computer science, data science, and business administration disciplines.", None),
    ("Risk Management Framework", 1),
    ("The enterprise risk committee identified cybersecurity, supply chain concentration, and "
     "regulatory divergence as the top three risk categories for the upcoming period. Mitigation "
     "strategies include a $12 million investment in security operations center capabilities, "
     "diversification of cloud infrastructure providers, and proactive engagement with regulatory "
     "bodies in key jurisdictions.", None),
    ("Business continuity planning was validated through two full-scale disaster recovery "
     "exercises, achieving recovery time objectives of under four hours for all Tier-1 systems. "
     "The company maintains cyber insurance coverage of $50 million and has implemented advanced "
     "threat detection powered by our own Meridian Intelligence Suite.", None),
    ("Strategic Outlook", 1),
    ("Looking ahead, Meridian Technologies is well-positioned to capitalize on secular growth "
     "trends in enterprise digital transformation. Our three-year strategic plan targets compound "
     "annual revenue growth of 18-22%, operating margin expansion to 30%, and entry into three "
     "new geographic markets. The leadership team remains confident in our ability to deliver "
     "sustained value creation for all stakeholders.", None),
    ("The board has authorized management to explore strategic partnerships in quantum computing "
     "readiness and edge computing infrastructure, recognizing these as critical capabilities for "
     "the next generation of enterprise workloads. Initial pilot programs are expected to launch "
     "in Q2, with commercialization timelines dependent on technology maturity assessments.", None),
]


def create_initial():
    doc = Document()

    # Set up single section with header and footer on ALL pages
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # IMPORTANT: different_first_page_header_footer must be False
    # so that all pages (including page 1) show the same header/footer.
    section.different_first_page_header_footer = False

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "Meridian Technologies — Annual Strategic Review 2025"
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in hp.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r0 = fp.add_run("Page ")
    r0.font.size = Pt(9)
    # PAGE field code
    r1 = fp.add_run()
    r1.font.size = Pt(9)
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)
    r2 = fp.add_run()
    r2.font.size = Pt(9)
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    r3.font.size = Pt(9)
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)
    r4 = fp.add_run(" of 10")
    r4.font.size = Pt(9)

    # Add content paragraphs
    for text, level in PARAGRAPHS:
        if level is not None:
            # It's a heading
            heading = doc.add_heading(text, level=level)
        else:
            para = doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(8)
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"

    # Add page breaks to ensure we have 10 pages of content
    # The above text produces roughly 4-5 pages; add filler sections to reach 10
    additional_sections = [
        ("Appendix A: Quarterly Revenue Breakdown", [
            "Q1 Revenue: $64.2M (Cloud: $38.5M, On-Premise: $15.2M, Services: $10.5M)",
            "Q2 Revenue: $69.8M (Cloud: $42.1M, On-Premise: $14.7M, Services: $13.0M)",
            "Q3 Revenue: $74.1M (Cloud: $45.8M, On-Premise: $14.3M, Services: $14.0M)",
            "Q4 Revenue: $79.3M (Cloud: $49.2M, On-Premise: $13.9M, Services: $16.2M)",
            "",
            "Cloud services represented 61% of total revenue in Q4, up from 60% in Q1, "
            "confirming the continued shift in our revenue mix toward recurring subscription "
            "models. On-premise license revenue declined 8.5% year-over-year but remains a "
            "significant contributor, particularly in regulated industries where cloud adoption "
            "faces compliance barriers.",
        ]),
        ("Appendix B: Key Performance Indicators", [
            "Customer Acquisition Cost (CAC): $24,300 (down 12% YoY)",
            "Customer Lifetime Value (CLV): $412,000 (up 18% YoY)",
            "CLV/CAC Ratio: 17.0x (industry benchmark: 10-15x)",
            "Net Revenue Retention: 118% (up from 112%)",
            "Annual Recurring Revenue (ARR): $215.6M (up 28% YoY)",
            "Gross Margin: 71.2% (up 1.4 percentage points)",
            "Rule of 40 Score: 51 (Revenue Growth 23% + FCF Margin 18%)",
            "",
            "These metrics demonstrate the fundamental health of our business model and the "
            "effectiveness of our go-to-market strategy. The improvement in CLV/CAC ratio "
            "reflects both higher deal values and reduced customer acquisition costs through "
            "our partner channel expansion program.",
        ]),
        ("Appendix C: Technology Infrastructure Summary", [
            "Global data centers: 12 (4 North America, 3 Europe, 3 Asia-Pacific, 2 Latin America)",
            "Total compute capacity: 48,000 vCPUs across hybrid cloud infrastructure",
            "Average platform uptime: 99.97% (SLA target: 99.95%)",
            "Mean time to recovery (MTTR): 12 minutes for Tier-1 incidents",
            "API transaction volume: 2.8 billion requests per month",
            "Data processed: 4.7 petabytes monthly across all client workloads",
            "",
            "Infrastructure investments during the period focused on edge computing nodes in "
            "15 metropolitan areas, enabling sub-10ms response times for latency-sensitive "
            "applications. The migration to Kubernetes-based orchestration for all production "
            "workloads was completed ahead of schedule, reducing infrastructure costs by 22% "
            "while improving deployment velocity by 3.5x.",
            "",
            "Our security posture was strengthened with the deployment of a next-generation "
            "SIEM platform, implementation of zero-trust network architecture across all "
            "offices, and achievement of SOC 2 Type II and ISO 27001 certifications for our "
            "newest data center facilities.",
        ]),
        ("Appendix D: Board of Directors and Executive Leadership", [
            "Board Chair: Dr. Katherine Whitfield, former CEO of Vertex Dynamics",
            "CEO: James R. Thornton, appointed 2019",
            "CFO: Priya Ramanathan, appointed 2021",
            "CTO: Dr. Alexander Volkov, appointed 2020",
            "COO: Lisa Marchetti, appointed 2022",
            "General Counsel: David Okafor, appointed 2018",
            "",
            "The board underwent planned succession during the period, with two new independent "
            "directors joining: Dr. Mei-Ling Chan, a recognized authority in AI governance and "
            "ethics, and Robert Castellano, former CFO of a Fortune 100 technology company. "
            "These appointments bring the board to nine members, with seven meeting independence "
            "criteria under NYSE listing standards.",
        ]),
    ]

    for title, lines in additional_sections:
        doc.add_page_break()
        heading = doc.add_heading(title, level=1)
        for line in lines:
            if line == "":
                doc.add_paragraph("")
            else:
                para = doc.add_paragraph(line)
                para.paragraph_format.space_after = Pt(6)
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
