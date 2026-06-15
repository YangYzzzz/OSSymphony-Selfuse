"""
Reward Script: Apply 'First Page' page style to page 1 (no header/footer),
               'Default Page Style' for pages 2+ (with headers/footers).
Task ID: writer_bs_060
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): different_first_page_header_footer enabled (titlePg present)
  Component 2 (0.25): First page header is empty (disabled)
  Component 3 (0.25): First page footer is empty (disabled)
  Component 4 (0.15): Default header/footer still have content for pages 2+
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_060'


def persist_app_state(domain: str):
    """Best-effort save of any unsaved GUI state."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: must have at least one section
    if len(doc.sections) == 0:
        print("CRITICAL: Document has no sections")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: different_first_page_header_footer is enabled (0.35 points)
    # In OOXML this corresponds to the <w:titlePg/> element in sectPr.
    # Initial state: False (no titlePg). Golden state: True.
    try:
        has_diff_first = section.different_first_page_header_footer
        if has_diff_first:
            print(f"PASS: Component 1 — different_first_page_header_footer is True (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 — different_first_page_header_footer is False, expected True")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: First page header is empty / has no text (0.25 points)
    # Initial state: no first page header distinction (linked/default).
    # Golden state: first page header exists but is empty (no header on page 1).
    try:
        if not section.different_first_page_header_footer:
            # If titlePg is not set, first page header is same as default — FAIL
            print("FAIL: Component 2 — No first page header distinction; cannot verify")
        else:
            fp_header = section.first_page_header
            fp_header_text = ""
            if fp_header and fp_header.paragraphs:
                fp_header_text = "".join(p.text for p in fp_header.paragraphs).strip()
            if fp_header_text == "":
                print(f"PASS: Component 2 — First page header is empty (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — First page header has text: {repr(fp_header_text)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: First page footer is empty / has no text (0.25 points)
    # Initial state: no first page footer distinction.
    # Golden state: first page footer exists but is empty (no footer on page 1).
    try:
        if not section.different_first_page_header_footer:
            print("FAIL: Component 3 — No first page footer distinction; cannot verify")
        else:
            fp_footer = section.first_page_footer
            fp_footer_text = ""
            if fp_footer and fp_footer.paragraphs:
                fp_footer_text = "".join(p.text for p in fp_footer.paragraphs).strip()
            if fp_footer_text == "":
                print(f"PASS: Component 3 — First page footer is empty (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — First page footer has text: {repr(fp_footer_text)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Default header and footer still have content for pages 2+ (0.15 points)
    # This checks that the default (non-first-page) header/footer are preserved.
    # Initial: default header/footer have content. Golden: same content preserved.
    # However, this component is ONLY scored when titlePg is set (Component 1 passes),
    # because without titlePg the default header/footer apply to ALL pages including page 1,
    # which is the initial (pre-task) state. This makes the compound check task-change-dependent.
    try:
        if not section.different_first_page_header_footer:
            print("FAIL: Component 4 — Without titlePg, default h/f on all pages is pre-task state")
        else:
            default_header = section.header
            default_footer = section.footer
            hdr_text = ""
            ftr_text = ""
            if default_header and default_header.paragraphs:
                hdr_text = "".join(p.text for p in default_header.paragraphs).strip()
            if default_footer and default_footer.paragraphs:
                ftr_text = "".join(p.text for p in default_footer.paragraphs).strip()

            if hdr_text and ftr_text:
                print(f"PASS: Component 4 — Default header ({repr(hdr_text[:40])}) and footer ({repr(ftr_text[:40])}) have content (0.15 pts)")
                total_score += 0.15
            elif hdr_text:
                print(f"PARTIAL: Component 4 — Default header has content but footer is empty")
                total_score += 0.075
            elif ftr_text:
                print(f"PARTIAL: Component 4 — Default footer has content but header is empty")
                total_score += 0.075
            else:
                print(f"FAIL: Component 4 — Default header and footer are both empty")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
