"""
Reward Script: Customer Newsletter Visual Structure
Task ID: writer_mktg_038
Domain: libreoffice_writer
Scoring:
  - Component 1: Newsletter header (22pt bold, centered)          — 0.15 pts
  - Component 2: Article titles in Georgia 16pt bold               — 0.35 pts
  - Component 3: Body text in Arial 11pt regular                   — 0.25 pts
  - Component 4: Horizontal rules between articles (4 separators) — 0.25 pts
  Total: 1.0
"""

import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_038'
FILE_PATH = '/home/user/Desktop/customer_newsletter_march.docx'

# Known article titles from task context
ARTICLE_TITLES = [
    'New Feature: Advanced Reporting',
    'Customer Spotlight: Meridian Health',
    'Upcoming Webinar: Q2 Product Roadmap',
    'Tips & Tricks: Keyboard Shortcuts',
    'Company News: Office Expansion',
]

HEADER_TEXT = 'The Insider'


def count_hr_paragraphs(doc):
    """Count paragraphs that have a bottom paragraph border (horizontal rule)."""
    count = 0
    hr_indices = []
    for i, para in enumerate(doc.paragraphs):
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                # Check for bottom border element
                bottom = pBdr.find(qn('w:bottom'))
                if bottom is not None:
                    count += 1
                    hr_indices.append(i)
    return count, hr_indices


def get_hr_space_before(doc, hr_indices):
    """Check space_before for separator paragraphs."""
    results = []
    for i in hr_indices:
        para = doc.paragraphs[i]
        pf = para.paragraph_format
        sp = pf.space_before.pt if pf.space_before else 0
        results.append(sp)
    return results


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    paragraphs = doc.paragraphs
    para_texts = [p.text.strip() for p in paragraphs]

    # Component 1: Newsletter header formatting — 22pt bold, centered (0.15 pts)
    # Precondition: must find header paragraph
    try:
        header_para = None
        for para in paragraphs:
            if HEADER_TEXT in para.text:
                header_para = para
                break

        if header_para is None:
            print(f"FAIL: Component 1 — Header paragraph containing '{HEADER_TEXT}' not found")
        else:
            runs = [r for r in header_para.runs if r.text.strip()]
            if not runs:
                print(f"FAIL: Component 1 — Header paragraph has no runs")
            else:
                run = runs[0]
                size_ok = run.font.size is not None and abs(run.font.size.pt - 22.0) < 0.5
                bold_ok = run.font.bold is True
                align_ok = (header_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)

                if size_ok and bold_ok and align_ok:
                    print(f"PASS: Component 1 — Header is 22pt bold centered (font={run.font.name}, size={run.font.size.pt}pt)")
                    total_score += 0.15
                else:
                    details = []
                    if not size_ok:
                        sz = run.font.size.pt if run.font.size else None
                        details.append(f"size={sz}pt (expected 22pt)")
                    if not bold_ok:
                        details.append(f"bold={run.font.bold} (expected True)")
                    if not align_ok:
                        details.append(f"alignment={header_para.paragraph_format.alignment} (expected CENTER)")
                    print(f"FAIL: Component 1 — Header not correctly formatted: {', '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Article titles in Georgia 16pt bold (0.35 pts)
    # Score = (number of correctly formatted titles) / 5 * 0.35
    try:
        title_scores = []
        for title in ARTICLE_TITLES:
            found = False
            for para in paragraphs:
                if para.text.strip() == title:
                    found = True
                    runs = [r for r in para.runs if r.text.strip()]
                    if not runs:
                        title_scores.append((title, False, "no runs"))
                        break
                    run = runs[0]
                    font_ok = run.font.name is not None and 'Georgia' in run.font.name
                    size_ok = run.font.size is not None and abs(run.font.size.pt - 16.0) < 0.5
                    bold_ok = run.font.bold is True
                    if font_ok and size_ok and bold_ok:
                        title_scores.append((title, True, f"Georgia {run.font.size.pt}pt bold"))
                    else:
                        issues = []
                        if not font_ok:
                            issues.append(f"font={run.font.name!r} (expected Georgia)")
                        if not size_ok:
                            sz = run.font.size.pt if run.font.size else None
                            issues.append(f"size={sz}pt (expected 16pt)")
                        if not bold_ok:
                            issues.append(f"bold={run.font.bold}")
                        title_scores.append((title, False, ", ".join(issues)))
                    break
            if not found:
                title_scores.append((title, False, "paragraph not found"))

        passing = sum(1 for _, ok, _ in title_scores if ok)
        component2_score = (passing / len(ARTICLE_TITLES)) * 0.35
        total_score += component2_score

        for title, ok, detail in title_scores:
            status = "PASS" if ok else "FAIL"
            print(f"  {status}: Title '{title[:40]}' — {detail}")
        print(f"{'PASS' if passing == len(ARTICLE_TITLES) else 'PARTIAL'}: Component 2 — {passing}/{len(ARTICLE_TITLES)} article titles in Georgia 16pt bold ({component2_score:.2f} pts)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Body text in Arial 11pt regular (0.25 pts)
    # Body paragraphs are those NOT matching a title, NOT the header, and NOT empty
    try:
        body_paras = []
        for para in paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if HEADER_TEXT in text:
                continue
            if text in ARTICLE_TITLES:
                continue
            body_paras.append(para)

        if not body_paras:
            print("FAIL: Component 3 — No body paragraphs found")
        else:
            correctly_formatted = 0
            total_body = len(body_paras)
            failed_examples = []

            for para in body_paras:
                runs = [r for r in para.runs if r.text.strip()]
                if not runs:
                    continue
                para_ok = True
                for run in runs:
                    font_ok = run.font.name is not None and (
                        'Arial' in run.font.name or 'Liberation Sans' in run.font.name
                    )
                    # Accept Arial OR common sans-serif fallbacks
                    size_ok = run.font.size is not None and abs(run.font.size.pt - 11.0) < 0.5
                    bold_ok = run.font.bold is not True  # not bold (None or False)
                    if not (font_ok and size_ok and bold_ok):
                        para_ok = False
                        if len(failed_examples) < 3:
                            failed_examples.append(
                                f"'{para.text[:30]}': font={run.font.name!r}, size={run.font.size.pt if run.font.size else None}pt, bold={run.font.bold}"
                            )
                        break
                if para_ok:
                    correctly_formatted += 1

            ratio = correctly_formatted / max(total_body, 1)
            component3_score = ratio * 0.25
            if component3_score > 0:
                total_score += component3_score

            if correctly_formatted == total_body:
                print(f"PASS: Component 3 — All {total_body} body paragraphs in Arial/sans-serif 11pt regular ({component3_score:.2f} pts)")
            else:
                print(f"PARTIAL: Component 3 — {correctly_formatted}/{total_body} body paragraphs correctly formatted ({component3_score:.2f} pts)")
                for ex in failed_examples:
                    print(f"  FAIL example: {ex}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Horizontal rules between articles (4 separators with space_before >= 12pt) (0.25 pts)
    # Horizontal rules implemented as empty paragraphs with bottom paragraph border
    try:
        hr_count, hr_indices = count_hr_paragraphs(doc)
        space_befores = get_hr_space_before(doc, hr_indices)

        # Need exactly 4 horizontal rules (between 5 articles)
        hr_count_ok = hr_count >= 4
        # All separators should have space_before >= 12pt
        spacing_ok = all(sp >= 12.0 for sp in space_befores) if space_befores else False

        if hr_count_ok and spacing_ok:
            print(f"PASS: Component 4 — {hr_count} horizontal rules found with space_before >= 12pt {space_befores}")
            total_score += 0.25
        elif hr_count_ok and not spacing_ok:
            # Partial: has HRs but spacing is wrong — still award partial
            qualifying = sum(1 for sp in space_befores if sp >= 12.0)
            partial = (qualifying / hr_count) * 0.25
            total_score += partial
            print(f"PARTIAL: Component 4 — {hr_count} HRs found, {qualifying} with space_before>=12pt. Space befores: {space_befores} ({partial:.2f} pts)")
        elif not hr_count_ok and hr_count > 0:
            # Has some HRs but not enough
            partial = min(hr_count / 4, 1.0) * 0.25
            total_score += partial
            print(f"PARTIAL: Component 4 — Only {hr_count}/4 horizontal rules found ({partial:.2f} pts)")
        else:
            print(f"FAIL: Component 4 — No horizontal rule paragraphs (pBdr/bottom border) found. hr_count={hr_count}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
