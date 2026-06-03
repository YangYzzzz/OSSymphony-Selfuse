"""
Initial Setup: Customer newsletter with no visual structure
Task ID: writer_mktg_038
Domain: libreoffice_writer

Creates a newsletter with 5 article sections, all in 12pt Liberation Sans,
no horizontal rules, no special formatting on titles.
File is placed at ~/Desktop/customer_newsletter_march.docx
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_038'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/customer_newsletter_march.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_run_font(run, font_name, size_pt, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set default font to Liberation Sans 12pt via Normal style
    style = doc.styles['Normal']
    style.font.name = 'Liberation Sans'
    style.font.size = Pt(12)

    # ---- Newsletter Header ----
    header_para = doc.add_paragraph()
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = header_para.add_run('The Insider \u2014 March 2026 Edition')
    set_run_font(run, 'Liberation Sans', 12, bold=False)

    # ---- Article 1: New Feature: Advanced Reporting ----
    title1 = doc.add_paragraph()
    r = title1.add_run('New Feature: Advanced Reporting')
    set_run_font(r, 'Liberation Sans', 12, bold=False)

    body1a = doc.add_paragraph()
    r = body1a.add_run(
        'We are excited to announce the launch of Advanced Reporting, a powerful new feature '
        'that gives you deeper insights into your business performance. With customizable '
        'dashboards and real-time data visualization, you can now make smarter decisions faster.'
    )
    set_run_font(r, 'Liberation Sans', 12)

    body1b = doc.add_paragraph()
    r = body1b.add_run(
        'Advanced Reporting is available to all Professional and Enterprise plan subscribers '
        'starting April 1, 2026. Log in to your dashboard today to explore the new tools.'
    )
    set_run_font(r, 'Liberation Sans', 12)

    # ---- Article 2: Customer Spotlight: Meridian Health ----
    title2 = doc.add_paragraph()
    r = title2.add_run('Customer Spotlight: Meridian Health')
    set_run_font(r, 'Liberation Sans', 12, bold=False)

    body2a = doc.add_paragraph()
    r = body2a.add_run(
        'This month we spotlight Meridian Health, a regional healthcare provider that has '
        'transformed its patient communication workflows using our platform. Since adopting '
        'the solution in late 2024, Meridian Health has reduced administrative overhead by 35%.'
    )
    set_run_font(r, 'Liberation Sans', 12)

    body2b = doc.add_paragraph()
    r = body2b.add_run(
        '"The integration with our existing systems was seamless," said Operations Director '
        'Laura Kinsella. "Our staff now spend more time on patient care and less on paperwork."'
    )
    set_run_font(r, 'Liberation Sans', 12)

    # ---- Article 3: Upcoming Webinar: Q2 Product Roadmap ----
    title3 = doc.add_paragraph()
    r = title3.add_run('Upcoming Webinar: Q2 Product Roadmap')
    set_run_font(r, 'Liberation Sans', 12, bold=False)

    body3a = doc.add_paragraph()
    r = body3a.add_run(
        'Join our product team on April 15, 2026 at 2:00 PM EST for an exclusive webinar '
        'covering the Q2 product roadmap. We will walk through planned enhancements, '
        'introduce beta features, and answer your questions live.'
    )
    set_run_font(r, 'Liberation Sans', 12)

    body3b = doc.add_paragraph()
    r = body3b.add_run(
        'Space is limited. Register now at insider.example.com/webinar-q2 to secure your spot. '
        'All registered attendees will receive a recording after the event.'
    )
    set_run_font(r, 'Liberation Sans', 12)

    # ---- Article 4: Tips & Tricks: Keyboard Shortcuts ----
    title4 = doc.add_paragraph()
    r = title4.add_run('Tips & Tricks: Keyboard Shortcuts')
    set_run_font(r, 'Liberation Sans', 12, bold=False)

    body4a = doc.add_paragraph()
    r = body4a.add_run(
        'Did you know that mastering keyboard shortcuts can cut your workflow time by up to 20%? '
        'This month we highlight some of our most popular shortcuts that power users rely on '
        'every day to navigate the platform efficiently.'
    )
    set_run_font(r, 'Liberation Sans', 12)

    body4b = doc.add_paragraph()
    r = body4b.add_run(
        'Try Ctrl+Shift+N to create a new project instantly, Ctrl+F to search across all records, '
        'and Alt+D to jump directly to the dashboard. Visit our Help Center for the full shortcut guide.'
    )
    set_run_font(r, 'Liberation Sans', 12)

    # ---- Article 5: Company News: Office Expansion ----
    title5 = doc.add_paragraph()
    r = title5.add_run('Company News: Office Expansion')
    set_run_font(r, 'Liberation Sans', 12, bold=False)

    body5a = doc.add_paragraph()
    r = body5a.add_run(
        'We are thrilled to announce the opening of our new Austin, Texas office this coming May. '
        'The expansion reflects our rapid growth in the US market and our commitment to providing '
        'local support to our growing customer base in the Southwest region.'
    )
    set_run_font(r, 'Liberation Sans', 12)

    body5b = doc.add_paragraph()
    r = body5b.add_run(
        'The Austin office will house our regional sales team and a dedicated customer success '
        'unit. If you are in the area, we would love to invite you to our open house event on May 20, 2026.'
    )
    set_run_font(r, 'Liberation Sans', 12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
