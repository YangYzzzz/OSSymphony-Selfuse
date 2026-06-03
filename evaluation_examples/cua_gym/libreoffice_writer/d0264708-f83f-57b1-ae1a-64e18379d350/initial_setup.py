"""
Initial Setup: Set fixed line spacing of 14pt for legal disclaimer paragraphs
Task ID: writer_para_024
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_para_024'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Title (Heading 1)
    doc.add_heading('END USER LICENSE AGREEMENT', level=1)

    # Paragraph 2: Date line
    doc.add_paragraph('Last updated: February 28, 2025')

    # Paragraph 3: Acceptance clause
    doc.add_paragraph(
        'By installing or using this software, you agree to be bound by the '
        'terms of this agreement.'
    )

    # Paragraph 4: Section heading — LIMITATION OF LIABILITY
    heading_lol = doc.add_paragraph('LIMITATION OF LIABILITY')
    heading_lol.runs[0].bold = True

    # Paragraph 5: Liability disclaimer text (NO fixed line spacing in initial)
    doc.add_paragraph(
        'IN NO EVENT SHALL THE LICENSOR BE LIABLE FOR ANY INDIRECT, INCIDENTAL, '
        'SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES INCLUDING BUT NOT LIMITED TO '
        'PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES, LOSS OF USE, DATA, OR PROFITS, '
        'OR BUSINESS INTERRUPTION.'
    )

    # Paragraph 6: Section heading — GOVERNING LAW
    heading_gl = doc.add_paragraph('GOVERNING LAW')
    heading_gl.runs[0].bold = True

    # Paragraph 7: Governing law text (NO fixed line spacing in initial)
    doc.add_paragraph(
        'This agreement shall be governed by and construed in accordance with the laws '
        'of the State of California, without regard to its conflict of laws provisions.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the initial artifact in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
