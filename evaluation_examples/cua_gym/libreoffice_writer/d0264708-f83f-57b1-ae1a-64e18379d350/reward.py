"""
Reward Script: Set a fixed line spacing of 14pt for the legal disclaimer text
Task ID: writer_para_024
Domain: libreoffice_writer
Scoring:
  Component 1: Para 4 (LIMITATION OF LIABILITY body) has FIXED line spacing at 14pt — 0.5 pts
  Component 2: Para 6 (GOVERNING LAW body) has FIXED line spacing at 14pt — 0.5 pts
Total: 1.0
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

WORKDIR = '/home/user'
TASK_ID = 'writer_para_024'

# 14pt in EMU: 1pt = 12700 EMU
TARGET_SPACING_EMU = int(14 * 12700)  # 177800
TOLERANCE_EMU = 100  # small tolerance for rounding


def verify_task(file_path):
    """
    Verify that legal disclaimer paragraphs (paragraphs 4 and 6, 0-indexed)
    have FIXED (EXACTLY) line spacing of 14pt.

    Para 4: 'IN NO EVENT SHALL THE LICENSOR...' (LIMITATION OF LIABILITY body)
    Para 6: 'This agreement shall be governed...' (GOVERNING LAW body)

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 7 paragraphs
    if len(doc.paragraphs) < 7:
        print(f"CRITICAL: Expected at least 7 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Paragraph 4 (LIMITATION OF LIABILITY body) has FIXED line spacing of 14pt (0.5 pts)
    try:
        para4 = doc.paragraphs[4]
        pf4 = para4.paragraph_format
        ls4 = pf4.line_spacing
        ls_rule4 = pf4.line_spacing_rule
        # Check that the text is correct paragraph (precondition validation)
        text_preview4 = para4.text[:30] if para4.text else ''
        # Verify FIXED (EXACTLY) rule and 14pt spacing
        if (ls_rule4 == WD_LINE_SPACING.EXACTLY
                and ls4 is not None
                and abs(ls4 - TARGET_SPACING_EMU) <= TOLERANCE_EMU):
            print(f"PASS: Component 1 — Para 4 has FIXED 14pt line spacing "
                  f"(ls_emu={ls4}, rule={ls_rule4}, text={text_preview4!r}) (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Para 4 expected FIXED 14pt line spacing, "
                  f"found ls_emu={ls4}, rule={ls_rule4}, text={text_preview4!r}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Paragraph 6 (GOVERNING LAW body) has FIXED line spacing of 14pt (0.5 pts)
    try:
        para6 = doc.paragraphs[6]
        pf6 = para6.paragraph_format
        ls6 = pf6.line_spacing
        ls_rule6 = pf6.line_spacing_rule
        text_preview6 = para6.text[:30] if para6.text else ''
        # Verify FIXED (EXACTLY) rule and 14pt spacing
        if (ls_rule6 == WD_LINE_SPACING.EXACTLY
                and ls6 is not None
                and abs(ls6 - TARGET_SPACING_EMU) <= TOLERANCE_EMU):
            print(f"PASS: Component 2 — Para 6 has FIXED 14pt line spacing "
                  f"(ls_emu={ls6}, rule={ls_rule6}, text={text_preview6!r}) (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 2 — Para 6 expected FIXED 14pt line spacing, "
                  f"found ls_emu={ls6}, rule={ls_rule6}, text={text_preview6!r}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in a given env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
