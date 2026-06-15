"""
Reward Script: Create nested table in cell C4 of project matrix
Task ID: writer_tm_029
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Cell C4 contains a nested table
  Component 2 (0.25): Nested table is 2x2
  Component 3 (0.40): Nested table cells contain Plan, Execute, Review, Close
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_029'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table (the outer table)
    if len(doc.tables) < 1:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    outer = doc.tables[0]

    # Precondition: outer table must be 4x6 (not scored, just gated)
    outer_rows = len(outer.rows)
    outer_cols = len(outer.columns)
    if outer_rows != 4 or outer_cols != 6:
        print(f"FAIL: Outer table dimensions changed — expected 4x6, found {outer_rows}x{outer_cols}")
        print("REWARD: 0.0")
        return 0.0

    # Get cell C4 = row index 3, col index 2
    try:
        cell_c4 = outer.cell(3, 2)
    except Exception as e:
        print(f"CRITICAL: Cannot access cell (3,2): {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Cell C4 contains a nested table (0.35 points)
    try:
        nested_count = len(cell_c4.tables)
        if nested_count >= 1:
            print(f"PASS: Component 1 — Cell C4 has {nested_count} nested table(s) (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 — Cell C4 has no nested tables")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Nested table is 2x2 (0.25 points)
    nested = None
    try:
        if len(cell_c4.tables) >= 1:
            nested = cell_c4.tables[0]
            n_rows = len(nested.rows)
            n_cols = len(nested.columns)
            if n_rows == 2 and n_cols == 2:
                print(f"PASS: Component 2 — Nested table is 2x2 (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — Nested table is {n_rows}x{n_cols}, expected 2x2")
        else:
            print("FAIL: Component 2 — No nested table to check dimensions")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Nested table cells contain correct text (0.40 points)
    # Expected: (0,0)=Plan, (0,1)=Execute, (1,0)=Review, (1,1)=Close
    try:
        if nested is not None and len(nested.rows) >= 2 and len(nested.columns) >= 2:
            expected = {
                (0, 0): 'Plan',
                (0, 1): 'Execute',
                (1, 0): 'Review',
                (1, 1): 'Close',
            }
            matches = 0
            for (r, c), exp_text in expected.items():
                actual = nested.cell(r, c).text.strip()
                if actual == exp_text:
                    matches += 1
                    print(f"  PASS: Nested ({r},{c}) = {repr(actual)}")
                else:
                    print(f"  FAIL: Nested ({r},{c}) = {repr(actual)}, expected {repr(exp_text)}")

            # Award proportional credit: 0.10 per correct cell
            cell_score = matches * 0.10
            if matches > 0:
                total_score += cell_score
            if matches == 4:
                print(f"PASS: Component 3 — All 4 nested cells correct (0.40 pts)")
            else:
                print(f"PARTIAL: Component 3 — {matches}/4 cells correct ({cell_score:.2f} pts)")
        else:
            print("FAIL: Component 3 — No valid nested table to check content")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
