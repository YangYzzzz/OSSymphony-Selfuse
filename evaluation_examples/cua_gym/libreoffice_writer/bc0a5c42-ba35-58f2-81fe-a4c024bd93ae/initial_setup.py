"""
Initial Setup: Create a project matrix document with a 4x6 table.
Task ID: writer_tm_029
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_029'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("Project Matrix - Q2 2025 Deliverables", level=1)

    # Brief intro paragraph
    intro = doc.add_paragraph(
        "This matrix tracks the key deliverables and responsibilities "
        "across departments for the second quarter of 2025. Each cell "
        "captures assigned owners, deadlines, and current status."
    )

    # Create 4x6 outer table
    # Columns: Task Area | Owner | Phase Info | Deadline | Budget ($K) | Status
    # Rows: Header + 3 data rows => 4 rows total
    table = doc.add_table(rows=4, cols=6)
    table.style = "Table Grid"

    # --- Row 0: Headers ---
    headers = ["Task Area", "Owner", "Phase Info", "Deadline", "Budget ($K)", "Status"]
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # --- Row 1: Data row ---
    row1_data = [
        "Cloud Migration",
        "Sarah Chen",
        "Infrastructure assessment and vendor selection for AWS migration",
        "2025-04-30",
        "185",
        "In Progress",
    ]
    for col_idx, val in enumerate(row1_data):
        table.cell(1, col_idx).text = val

    # --- Row 2: Data row ---
    row2_data = [
        "CRM Integration",
        "Marcus Johnson",
        "API development and data mapping for Salesforce integration",
        "2025-05-15",
        "120",
        "Planning",
    ]
    for col_idx, val in enumerate(row2_data):
        table.cell(2, col_idx).text = val

    # --- Row 3: Data row (C4 is cell(3, 2) -- contains "Phase Details") ---
    row3_data = [
        "Security Audit",
        "Priya Kapoor",
        "Phase Details",
        "2025-06-20",
        "95",
        "Not Started",
    ]
    for col_idx, val in enumerate(row3_data):
        table.cell(3, col_idx).text = val

    # Additional context paragraph after the table
    doc.add_paragraph("")
    note = doc.add_paragraph(
        "Note: Budget figures are preliminary estimates subject to quarterly review. "
        "Contact the PMO office for updated allocations."
    )

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
