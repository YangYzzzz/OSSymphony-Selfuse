"""
Initial Setup: Create a 30-page legal contract with section headings, bookmarks,
and 12 manual text references like 'see Section 3.2' (plain text, not linked).
Task ID: writer_legal_053
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_053'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_bookmark(paragraph, bookmark_name, bookmark_id):
    """Add a bookmark spanning the entire paragraph text."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    tag_start = OxmlElement('w:bookmarkStart')
    tag_start.set(qn('w:id'), str(bookmark_id))
    tag_start.set(qn('w:name'), bookmark_name)
    tag_end = OxmlElement('w:bookmarkEnd')
    tag_end.set(qn('w:id'), str(bookmark_id))
    # Insert bookmark start before the first run, end after last run
    paragraph._element.insert(0, tag_start)
    paragraph._element.append(tag_end)


def set_heading_style(para, level=2):
    """Apply heading style with consistent formatting."""
    para.style = f'Heading {level}'
    for run in para.runs:
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.bold = True


def add_body_text(doc, text, space_after=Pt(6)):
    """Add a body paragraph with standard formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    para.paragraph_format.space_after = space_after
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return para


# Define the contract structure: (section_number, heading_text)
SECTIONS = [
    ("1", "Definitions and Interpretation"),
    ("1.1", "Key Definitions"),
    ("1.2", "Rules of Interpretation"),
    ("1.3", "Order of Precedence"),
    ("2", "Scope of Services"),
    ("2.1", "Primary Service Obligations"),
    ("2.2", "Service Level Requirements"),
    ("2.3", "Excluded Services"),
    ("3", "Payment Terms"),
    ("3.1", "Fee Schedule"),
    ("3.2", "Invoicing Procedures"),
    ("3.3", "Late Payment Penalties"),
    ("3.4", "Currency and Exchange Rates"),
    ("4", "Term and Termination"),
    ("4.1", "Initial Term"),
    ("4.2", "Renewal Provisions"),
    ("4.3", "Termination for Cause"),
    ("4.4", "Termination for Convenience"),
    ("5", "Intellectual Property"),
    ("5.1", "Ownership of Pre-Existing IP"),
    ("5.2", "License Grants"),
    ("5.3", "Work Product Assignment"),
    ("6", "Confidentiality"),
    ("6.1", "Definition of Confidential Information"),
    ("6.2", "Obligations of Receiving Party"),
    ("6.3", "Permitted Disclosures"),
    ("7", "Representations and Warranties"),
    ("7.1", "Mutual Representations"),
    ("7.2", "Service Provider Warranties"),
    ("7.3", "Client Warranties"),
    ("8", "Limitation of Liability"),
    ("8.1", "Cap on Liability"),
    ("8.2", "Exclusion of Consequential Damages"),
    ("8.3", "Carve-Outs"),
    ("9", "Indemnification"),
    ("9.1", "Service Provider Indemnification"),
    ("9.2", "Client Indemnification"),
    ("9.3", "Indemnification Procedures"),
    ("10", "Dispute Resolution"),
    ("10.1", "Negotiation"),
    ("10.2", "Mediation"),
    ("10.3", "Arbitration"),
    ("11", "General Provisions"),
    ("11.1", "Force Majeure"),
    ("11.2", "Assignment"),
    ("11.3", "Notices"),
    ("11.4", "Entire Agreement"),
    ("11.5", "Amendments"),
    ("11.6", "Governing Law"),
]

# 12 cross-references: (appears_in_section_index, reference_text, target_section_number)
# These are the 12 manual text references scattered throughout the document
CROSS_REFS = [
    (3, "as defined in Section 1.1", "1.1"),
    (6, "subject to the payment terms set forth in Section 3.1", "3.1"),
    (8, "in accordance with Section 2.2", "2.2"),
    (12, "except as provided in Section 4.3", "4.3"),
    (16, "pursuant to Section 6.1", "6.1"),
    (19, "as further described in Section 5.2", "5.2"),
    (22, "see Section 7.2 for applicable warranties", "7.2"),
    (25, "subject to the limitations in Section 8.1", "8.1"),
    (29, "the procedures outlined in Section 9.3", "9.3"),
    (33, "in accordance with Section 10.2", "10.2"),
    (37, "as set forth in Section 3.2", "3.2"),
    (42, "subject to Section 11.4", "11.4"),
]

# Map section numbers to bookmark names
def section_to_bookmark(section_num):
    return f"_Section_{section_num.replace('.', '_')}"


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # -- Title page --
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("MASTER SERVICES AGREEMENT")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Times New Roman'

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Between Nexus Technologies Inc. and Meridian Solutions Ltd.")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.space_before = Pt(24)
    run = date_para.add_run("Effective Date: January 15, 2026")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_page_break()

    # -- Table of Contents placeholder --
    toc_heading = doc.add_heading("Table of Contents", level=1)
    for sec_num, sec_title in SECTIONS:
        level = sec_num.count('.')
        indent = "    " * level
        toc_entry = doc.add_paragraph(f"{indent}{sec_num}  {sec_title}")
        toc_entry.paragraph_format.space_after = Pt(2)
        for run in toc_entry.runs:
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

    doc.add_page_break()

    # -- Preamble --
    add_body_text(doc,
        'This Master Services Agreement ("Agreement") is entered into as of January 15, 2026 '
        '("Effective Date"), by and between Nexus Technologies Inc., a Delaware corporation with '
        'its principal place of business at 4200 Innovation Drive, Suite 800, San Francisco, '
        'California 94105 ("Client"), and Meridian Solutions Ltd., a corporation organized under '
        'the laws of the State of New York with its principal place of business at 1750 Broadway, '
        '22nd Floor, New York, New York 10019 ("Service Provider").')

    add_body_text(doc,
        "WHEREAS, Client desires to engage Service Provider to provide certain professional "
        "services, technology consulting, and related deliverables as more fully described herein; and")

    add_body_text(doc,
        "WHEREAS, Service Provider has the expertise, personnel, and resources necessary to "
        "perform such services and desires to provide the same to Client on the terms and "
        "conditions set forth in this Agreement;")

    add_body_text(doc,
        "NOW, THEREFORE, in consideration of the mutual covenants, promises, and agreements "
        "contained herein, and for other good and valuable consideration, the receipt and "
        "sufficiency of which are hereby acknowledged, the parties agree as follows:")

    # Track which sections contain cross-references
    cross_ref_map = {}
    for sec_idx, ref_text, target in CROSS_REFS:
        cross_ref_map.setdefault(sec_idx, []).append((ref_text, target))

    # Body content for each section (realistic legal text)
    SECTION_BODY = {
        "1": [
            'This Article sets forth the definitions and interpretive rules that govern this Agreement. '
            'All capitalized terms used in this Agreement shall have the meanings ascribed to them in this Article '
            'or as otherwise defined throughout this Agreement.'
        ],
        "1.1": [
            '"Affiliate" means any entity that directly or indirectly controls, is controlled by, or is under '
            'common control with a party, where "control" means ownership of fifty percent (50%) or more of the '
            'outstanding voting securities of such entity.',
            '"Deliverables" means all work product, software, documentation, reports, analyses, and other materials '
            'created by Service Provider in the performance of the Services under this Agreement.',
            '"Intellectual Property Rights" means all patents, copyrights, trademarks, trade secrets, and other '
            'intellectual property rights recognized in any jurisdiction worldwide.',
            '"Services" means the professional services, technology consulting, software development, and related '
            'services to be performed by Service Provider as described in this Agreement and applicable Statements of Work.',
        ],
        "1.2": [
            'In this Agreement, unless the context otherwise requires: (a) words importing the singular include '
            'the plural and vice versa; (b) words importing a gender include every gender; (c) references to '
            '"including" or "includes" mean including without limitation; (d) references to statutes or statutory '
            'provisions shall include any statute or statutory provision which amends, extends, consolidates, or '
            'replaces the same.',
        ],
        "1.3": [
            'In the event of any conflict or inconsistency between the provisions of this Agreement and any '
            'Statement of Work, the provisions of this Agreement shall prevail unless the Statement of Work '
            'expressly states that it is intended to override a specific provision of this Agreement.',
        ],
        "2": [
            'Service Provider shall provide to Client the Services described in this Article and in the applicable '
            'Statements of Work executed by both parties from time to time during the Term of this Agreement.',
        ],
        "2.1": [
            'Service Provider shall perform the Services with the degree of skill, care, and diligence that would '
            'reasonably be expected of a qualified professional in the same field. Service Provider shall assign '
            'qualified personnel with appropriate experience and expertise to perform the Services.',
            'All Services shall be performed in accordance with the specifications, requirements, and timelines '
            'set forth in the applicable Statement of Work. Service Provider shall promptly notify Client of any '
            'circumstances that may delay or materially affect the performance of the Services.',
        ],
        "2.2": [
            'Service Provider shall meet or exceed the service levels specified in Schedule B attached hereto. '
            'Service levels shall be measured on a monthly basis, and Service Provider shall provide written '
            'reports to Client detailing its performance against each service level metric.',
            'In the event that Service Provider fails to meet any service level for two (2) consecutive months, '
            'Client shall be entitled to the service credits described in Schedule B, without prejudice to any '
            'other remedies available under this Agreement.',
        ],
        "2.3": [
            'The following services are expressly excluded from the scope of this Agreement unless otherwise '
            'agreed in writing by the parties: (a) hardware procurement and installation; (b) third-party software '
            'licensing; (c) data migration from legacy systems not identified in a Statement of Work; and '
            '(d) end-user training beyond what is specified in the applicable Statement of Work.',
        ],
        "3": [
            'Client shall pay Service Provider for the Services in accordance with the terms set forth in this '
            'Article. All fees are exclusive of applicable taxes, which shall be the responsibility of Client '
            'unless otherwise specified.',
        ],
        "3.1": [
            'The fees for the Services shall be as set forth in Schedule A attached hereto. Service Provider '
            'shall not increase fees during the Initial Term without Client\'s prior written consent. For any '
            'renewal term, Service Provider may increase fees by providing at least ninety (90) days\' written '
            'notice prior to the commencement of such renewal term, provided that any increase shall not exceed '
            'five percent (5%) of the then-current fees.',
        ],
        "3.2": [
            'Service Provider shall submit invoices to Client on a monthly basis, no later than the tenth (10th) '
            'business day following the end of each calendar month. Each invoice shall include a detailed breakdown '
            'of Services performed, hours expended (for time-and-materials engagements), expenses incurred, and '
            'any applicable taxes.',
            'Client shall pay each undisputed invoice within thirty (30) days of receipt. Client may dispute any '
            'portion of an invoice in good faith by providing written notice to Service Provider within fifteen (15) '
            'days of receipt, specifying the nature and basis of the dispute.',
        ],
        "3.3": [
            'If Client fails to pay any undisputed amount when due, Service Provider shall be entitled to charge '
            'interest on the overdue amount at the rate of one and one-half percent (1.5%) per month, or the maximum '
            'rate permitted by applicable law, whichever is less, from the date payment was due until the date of '
            'actual payment.',
        ],
        "3.4": [
            'All payments under this Agreement shall be made in United States Dollars (USD). If Service Provider '
            'incurs expenses in a foreign currency, such expenses shall be converted to USD using the exchange rate '
            'published by the Federal Reserve Bank of New York on the date the expense was incurred.',
        ],
        "4": [
            'This Article governs the duration of this Agreement, the conditions under which it may be renewed, '
            'and the circumstances under which either party may terminate it.',
        ],
        "4.1": [
            'This Agreement shall commence on the Effective Date and shall continue for an initial period of '
            'thirty-six (36) months ("Initial Term"), unless earlier terminated in accordance with the provisions '
            'of this Article.',
        ],
        "4.2": [
            'Upon expiration of the Initial Term, this Agreement shall automatically renew for successive twelve '
            '(12) month periods ("Renewal Terms"), unless either party provides written notice of non-renewal at '
            'least ninety (90) days prior to the expiration of the then-current term.',
        ],
        "4.3": [
            'Either party may terminate this Agreement immediately upon written notice if the other party: '
            '(a) commits a material breach of this Agreement and fails to cure such breach within thirty (30) days '
            'after receiving written notice thereof; (b) becomes insolvent, files a petition in bankruptcy, or '
            'makes an assignment for the benefit of creditors; or (c) ceases to conduct business in the normal course.',
        ],
        "4.4": [
            'Either party may terminate this Agreement for convenience upon one hundred twenty (120) days\' prior '
            'written notice to the other party. In the event of termination for convenience by Client, Client shall '
            'pay Service Provider for all Services performed and expenses incurred through the effective date of '
            'termination, plus any reasonable wind-down costs approved in advance by Client.',
        ],
        "5": [
            'This Article addresses the ownership, licensing, and assignment of intellectual property created '
            'or used in connection with the Services.',
        ],
        "5.1": [
            'Each party shall retain all right, title, and interest in and to its Pre-Existing IP. "Pre-Existing '
            'IP" means all Intellectual Property Rights owned or controlled by a party prior to the Effective Date '
            'or developed by a party independently of this Agreement.',
        ],
        "5.2": [
            'Service Provider hereby grants to Client a non-exclusive, worldwide, royalty-free, perpetual license '
            'to use any Pre-Existing IP of Service Provider that is incorporated into or necessary for the use '
            'of the Deliverables, solely in connection with Client\'s use of such Deliverables.',
            'Client hereby grants to Service Provider a non-exclusive, limited license to use Client\'s Pre-Existing '
            'IP solely to the extent necessary for Service Provider to perform the Services during the Term.',
        ],
        "5.3": [
            'All Deliverables created by Service Provider specifically for Client under this Agreement shall be '
            'considered "work made for hire" to the maximum extent permitted by applicable law. To the extent any '
            'Deliverable does not qualify as work made for hire, Service Provider hereby irrevocably assigns to '
            'Client all right, title, and interest in and to such Deliverable.',
        ],
        "6": [
            'The parties acknowledge that in the course of performing their obligations under this Agreement, each '
            'party may disclose or make available to the other party certain confidential and proprietary information.',
        ],
        "6.1": [
            '"Confidential Information" means any and all non-public information disclosed by one party (the '
            '"Disclosing Party") to the other party (the "Receiving Party"), whether orally, in writing, or by '
            'inspection, that is designated as confidential or that a reasonable person would understand to be '
            'confidential given the nature of the information and the circumstances of disclosure.',
            'Confidential Information includes, without limitation: business plans, financial data, customer lists, '
            'pricing strategies, technical specifications, source code, algorithms, trade secrets, and any other '
            'information that derives independent economic value from not being generally known.',
        ],
        "6.2": [
            'The Receiving Party shall: (a) hold the Disclosing Party\'s Confidential Information in strict '
            'confidence; (b) not disclose such information to any third party without the Disclosing Party\'s '
            'prior written consent; (c) use such information solely for the purposes of this Agreement; and '
            '(d) protect such information with the same degree of care it uses to protect its own confidential '
            'information, but in no event less than reasonable care.',
        ],
        "6.3": [
            'The confidentiality obligations set forth herein shall not apply to information that: (a) is or '
            'becomes publicly available through no fault of the Receiving Party; (b) was lawfully in the '
            'Receiving Party\'s possession prior to disclosure; (c) is independently developed by the Receiving '
            'Party without reference to the Confidential Information; or (d) is required to be disclosed by law, '
            'regulation, or court order, provided that the Receiving Party gives the Disclosing Party prompt '
            'written notice of such requirement.',
        ],
        "7": [
            'Each party makes the following representations and warranties to the other party as of the Effective Date '
            'and throughout the Term of this Agreement.',
        ],
        "7.1": [
            'Each party represents and warrants that: (a) it is duly organized, validly existing, and in good '
            'standing under the laws of its jurisdiction of organization; (b) it has full corporate power and '
            'authority to enter into this Agreement and to perform its obligations hereunder; and (c) this '
            'Agreement constitutes a legal, valid, and binding obligation enforceable against it in accordance '
            'with its terms.',
        ],
        "7.2": [
            'Service Provider represents and warrants that: (a) the Services shall be performed in a professional '
            'and workmanlike manner by qualified personnel; (b) the Deliverables shall conform to the specifications '
            'set forth in the applicable Statement of Work; (c) the Services and Deliverables shall not infringe '
            'upon any third-party Intellectual Property Rights; and (d) Service Provider shall comply with all '
            'applicable laws and regulations in the performance of the Services.',
        ],
        "7.3": [
            'Client represents and warrants that: (a) it shall provide Service Provider with timely access to all '
            'information, systems, and resources reasonably necessary for Service Provider to perform the Services; '
            '(b) all information and materials provided by Client to Service Provider shall be accurate and complete; '
            'and (c) Client has obtained all necessary consents and authorizations for the processing of data '
            'provided to Service Provider.',
        ],
        "8": [
            'This Article sets forth the limitations on each party\'s liability under this Agreement.',
        ],
        "8.1": [
            'EXCEPT FOR THE OBLIGATIONS SET FORTH IN ARTICLES 6 (CONFIDENTIALITY) AND 9 (INDEMNIFICATION), '
            'NEITHER PARTY\'S AGGREGATE LIABILITY UNDER THIS AGREEMENT SHALL EXCEED THE TOTAL FEES PAID OR PAYABLE '
            'BY CLIENT TO SERVICE PROVIDER DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE EVENT '
            'GIVING RISE TO SUCH LIABILITY.',
        ],
        "8.2": [
            'IN NO EVENT SHALL EITHER PARTY BE LIABLE TO THE OTHER PARTY FOR ANY INDIRECT, INCIDENTAL, SPECIAL, '
            'CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING BUT NOT LIMITED TO LOSS OF PROFITS, LOSS OF REVENUE, '
            'LOSS OF DATA, OR BUSINESS INTERRUPTION, REGARDLESS OF THE THEORY OF LIABILITY AND EVEN IF SUCH PARTY '
            'HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.',
        ],
        "8.3": [
            'The limitations set forth in this Article shall not apply to: (a) damages arising from a party\'s '
            'willful misconduct or gross negligence; (b) damages arising from a breach of Article 6 (Confidentiality); '
            '(c) Service Provider\'s indemnification obligations under Section 9.1; or (d) damages arising from '
            'infringement of Intellectual Property Rights.',
        ],
        "9": [
            'This Article establishes the indemnification obligations of each party and the procedures for '
            'seeking indemnification.',
        ],
        "9.1": [
            'Service Provider shall defend, indemnify, and hold harmless Client and its Affiliates, directors, '
            'officers, employees, and agents from and against any and all claims, damages, losses, liabilities, '
            'costs, and expenses (including reasonable attorneys\' fees) arising from or related to: (a) any '
            'breach of Service Provider\'s representations, warranties, or obligations under this Agreement; '
            '(b) any claim that the Services or Deliverables infringe upon third-party Intellectual Property Rights; '
            'or (c) any negligent or wrongful act or omission of Service Provider or its personnel.',
        ],
        "9.2": [
            'Client shall defend, indemnify, and hold harmless Service Provider and its Affiliates, directors, '
            'officers, employees, and agents from and against any and all claims, damages, losses, liabilities, '
            'costs, and expenses (including reasonable attorneys\' fees) arising from or related to: (a) any '
            'breach of Client\'s representations, warranties, or obligations under this Agreement; (b) any '
            'materials or data provided by Client that infringe upon third-party rights; or (c) Client\'s use '
            'of the Deliverables in a manner not authorized by this Agreement.',
        ],
        "9.3": [
            'A party seeking indemnification ("Indemnified Party") shall: (a) promptly notify the indemnifying '
            'party ("Indemnifying Party") in writing of any claim for which indemnification is sought; (b) grant '
            'the Indemnifying Party sole control of the defense and settlement of such claim; and (c) provide '
            'reasonable cooperation and assistance to the Indemnifying Party at the Indemnifying Party\'s expense.',
            'The Indemnifying Party shall not settle any claim in a manner that imposes any obligation, restriction, '
            'or liability on the Indemnified Party without the Indemnified Party\'s prior written consent, which '
            'shall not be unreasonably withheld.',
        ],
        "10": [
            'The parties agree to resolve any disputes arising under or in connection with this Agreement in '
            'accordance with the procedures set forth in this Article.',
        ],
        "10.1": [
            'The parties shall first attempt to resolve any dispute through good faith negotiations between senior '
            'executives of each party. Either party may initiate negotiations by delivering written notice to the '
            'other party describing the dispute in reasonable detail. The parties shall use commercially reasonable '
            'efforts to resolve the dispute within thirty (30) days of such notice.',
        ],
        "10.2": [
            'If the parties are unable to resolve a dispute through negotiation within the period specified above, '
            'either party may submit the dispute to mediation administered by JAMS under its Mediation Rules. '
            'The mediation shall be conducted by a single mediator selected by mutual agreement of the parties, '
            'or, failing such agreement, appointed by JAMS. The mediation shall take place in New York, New York.',
        ],
        "10.3": [
            'If mediation fails to resolve the dispute within sixty (60) days of the commencement of mediation, '
            'either party may submit the dispute to final and binding arbitration administered by the American '
            'Arbitration Association ("AAA") under its Commercial Arbitration Rules. The arbitration shall be '
            'conducted by a panel of three (3) arbitrators, with each party selecting one arbitrator and the two '
            'party-appointed arbitrators selecting the third.',
            'The arbitration shall take place in New York, New York, and the language of the arbitration shall '
            'be English. The arbitrator(s) shall render a reasoned award within ninety (90) days of the close of '
            'hearings. Judgment on the award may be entered in any court of competent jurisdiction.',
        ],
        "11": [
            'This Article contains miscellaneous provisions governing the administration and interpretation '
            'of this Agreement.',
        ],
        "11.1": [
            'Neither party shall be liable for any failure or delay in performing its obligations under this '
            'Agreement (other than payment obligations) to the extent such failure or delay results from '
            'circumstances beyond the affected party\'s reasonable control, including but not limited to acts '
            'of God, natural disasters, pandemics, war, terrorism, government actions, labor disputes, '
            'power failures, or internet or telecommunications failures.',
        ],
        "11.2": [
            'Neither party may assign or transfer this Agreement or any of its rights or obligations hereunder '
            'without the prior written consent of the other party, except that either party may assign this '
            'Agreement to an Affiliate or in connection with a merger, acquisition, or sale of all or substantially '
            'all of its assets, provided that the assignee agrees in writing to be bound by the terms of this Agreement.',
        ],
        "11.3": [
            'All notices and other communications required or permitted under this Agreement shall be in writing '
            'and shall be deemed duly given when: (a) delivered personally; (b) sent by certified mail, return '
            'receipt requested, postage prepaid; (c) sent by a nationally recognized overnight courier service; '
            'or (d) sent by email with confirmation of receipt, to the addresses set forth on the signature page '
            'or to such other address as a party may designate by written notice.',
        ],
        "11.4": [
            'This Agreement, together with all Schedules, Exhibits, and Statements of Work attached hereto or '
            'incorporated by reference, constitutes the entire agreement between the parties with respect to the '
            'subject matter hereof and supersedes all prior and contemporaneous agreements, understandings, '
            'negotiations, and discussions, whether oral or written, between the parties.',
        ],
        "11.5": [
            'No amendment, modification, or waiver of any provision of this Agreement shall be effective unless '
            'made in writing and signed by both parties. No failure or delay by either party in exercising any '
            'right or remedy under this Agreement shall operate as a waiver thereof, nor shall any single or '
            'partial exercise of any right or remedy preclude any other or further exercise thereof.',
        ],
        "11.6": [
            'This Agreement shall be governed by and construed in accordance with the laws of the State of '
            'New York, without giving effect to any principles of conflicts of law. The parties irrevocably '
            'submit to the exclusive jurisdiction of the state and federal courts located in the County of '
            'New York, State of New York, for the resolution of any disputes not subject to arbitration under '
            'Article 10.',
        ],
    }

    bookmark_id = 1

    for sec_idx, (sec_num, sec_title) in enumerate(SECTIONS):
        # Determine heading level
        level = 1 if '.' not in sec_num else 2

        # Add heading with section number
        heading_text = f"Section {sec_num}: {sec_title}"
        heading = doc.add_heading(heading_text, level=level)

        # Add bookmark on the heading
        bm_name = section_to_bookmark(sec_num)
        add_bookmark(heading, bm_name, bookmark_id)
        bookmark_id += 1

        # Add body text
        body_paragraphs = SECTION_BODY.get(sec_num, [
            f"The provisions of this section shall be interpreted in accordance with the general "
            f"terms and conditions of this Agreement."
        ])

        for body_text in body_paragraphs:
            add_body_text(doc, body_text)

        # Add cross-references as plain text (NOT linked)
        if sec_idx in cross_ref_map:
            for ref_text, target_sec in cross_ref_map[sec_idx]:
                ref_para = add_body_text(doc,
                    f"For additional details, {ref_text} of this Agreement.")

        # Add extra filler paragraphs to reach ~30 pages
        if level == 1:
            add_body_text(doc, "")  # blank separator

    # Signature block
    doc.add_page_break()
    sig_heading = doc.add_heading("Signatures", level=1)

    add_body_text(doc,
        "IN WITNESS WHEREOF, the parties hereto have caused this Agreement to be executed "
        "by their duly authorized representatives as of the Effective Date.")

    for _ in range(2):
        add_body_text(doc, "")

    # Nexus Technologies signature block
    add_body_text(doc, "NEXUS TECHNOLOGIES INC.")
    add_body_text(doc, "")
    add_body_text(doc, "By: ____________________________")
    add_body_text(doc, "Name: Dr. Alexandra Whitfield")
    add_body_text(doc, "Title: Chief Executive Officer")
    add_body_text(doc, "Date: January 15, 2026")

    for _ in range(2):
        add_body_text(doc, "")

    # Meridian Solutions signature block
    add_body_text(doc, "MERIDIAN SOLUTIONS LTD.")
    add_body_text(doc, "")
    add_body_text(doc, "By: ____________________________")
    add_body_text(doc, "Name: Jonathan R. Blackwell")
    add_body_text(doc, "Title: Managing Director")
    add_body_text(doc, "Date: January 15, 2026")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
