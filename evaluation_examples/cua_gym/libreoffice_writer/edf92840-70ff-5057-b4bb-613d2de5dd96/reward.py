"""
Reward Script: Avery 5160 label sheet with product SKU and name
Task ID: writer_lec_065
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Table exists with 10 rows x 3 cols (30 labels)
  Component 2 (0.35): All labels have SKU 'PRD-2025-001' in bold monospace font
  Component 3 (0.35): All labels have 'Wireless Bluetooth Headphones' in regular proportional font
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_065'
EXPECTED_SKU = 'PRD-2025-001'
EXPECTED_NAME = 'Wireless Bluetooth Headphones'
EXPECTED_ROWS = 10
EXPECTED_COLS = 3
EXPECTED_LABELS = 30

MONOSPACE_FONTS = {'courier new', 'courier', 'consolas', 'lucida console',
                   'dejavu sans mono', 'liberation mono', 'andale mono',
                   'monaco', 'menlo', 'source code pro', 'roboto mono',
                   'noto mono', 'ubuntu mono', 'fira mono', 'fira code'}


def is_monospace(font_name):
    """Check if font name is a known monospace font."""
    if not font_name:
        return False
    return font_name.lower().strip() in MONOSPACE_FONTS


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with correct dimensions (0.30 points)
    # Initial file has no tables, so this only passes on golden.
    try:
        if len(doc.tables) == 0:
            print(f"FAIL: Component 1 — No tables found in document")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            total_cells = num_rows * num_cols
            if num_cols == EXPECTED_COLS and total_cells >= EXPECTED_LABELS:
                print(f"PASS: Component 1 — Table has {num_rows} rows x {num_cols} cols = {total_cells} labels (0.30 pts)")
                total_score += 0.30
            else:
                print(f"FAIL: Component 1 — Table is {num_rows}x{num_cols}={total_cells} labels, expected at least {EXPECTED_LABELS} labels in 3-column layout")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Early exit if no table
    if len(doc.tables) == 0:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    table = doc.tables[0]

    # Component 2: All labels have SKU in bold monospace (0.35 points)
    # Check each cell for the SKU text on line 1, bold, monospace font
    try:
        sku_correct_count = 0
        total_cells_checked = 0
        for row in table.rows:
            for cell in row.cells:
                total_cells_checked += 1
                cell_runs = []
                for para in cell.paragraphs:
                    for run in para.runs:
                        cell_runs.append(run)

                if len(cell_runs) < 1:
                    continue

                first_run = cell_runs[0]
                has_sku = EXPECTED_SKU in first_run.text
                is_bold = first_run.font.bold is True
                is_mono = is_monospace(first_run.font.name)

                if has_sku and is_bold and is_mono:
                    sku_correct_count += 1

        if total_cells_checked > 0 and sku_correct_count >= EXPECTED_LABELS:
            print(f"PASS: Component 2 — {sku_correct_count}/{total_cells_checked} labels have bold monospace SKU (0.35 pts)")
            total_score += 0.35
        elif sku_correct_count > 0:
            partial = 0.35 * (sku_correct_count / EXPECTED_LABELS)
            print(f"PARTIAL: Component 2 — {sku_correct_count}/{total_cells_checked} labels correct, partial credit: {partial:.3f}")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — {sku_correct_count}/{total_cells_checked} labels have bold monospace SKU")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All labels have product name in regular proportional font (0.35 points)
    try:
        name_correct_count = 0
        total_cells_checked = 0
        for row in table.rows:
            for cell in row.cells:
                total_cells_checked += 1
                cell_runs = []
                for para in cell.paragraphs:
                    for run in para.runs:
                        cell_runs.append(run)

                if len(cell_runs) < 2:
                    continue

                second_run = cell_runs[1]
                has_name = EXPECTED_NAME in second_run.text
                is_not_bold = second_run.font.bold is not True  # False or None both acceptable
                is_not_mono = not is_monospace(second_run.font.name)

                if has_name and is_not_bold and is_not_mono:
                    name_correct_count += 1

        if total_cells_checked > 0 and name_correct_count >= EXPECTED_LABELS:
            print(f"PASS: Component 3 — {name_correct_count}/{total_cells_checked} labels have regular proportional name (0.35 pts)")
            total_score += 0.35
        elif name_correct_count > 0:
            partial = 0.35 * (name_correct_count / EXPECTED_LABELS)
            print(f"PARTIAL: Component 3 — {name_correct_count}/{total_cells_checked} labels correct, partial credit: {partial:.3f}")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — {name_correct_count}/{total_cells_checked} labels have regular proportional name")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
