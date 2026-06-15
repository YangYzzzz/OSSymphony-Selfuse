"""
Reward Script: Apply 'Emphasis' character style to species names in bio_notes.docx
Task ID: writer_txtfmt_054
Domain: libreoffice_writer
Scoring:
  Component 1: 'Homo sapiens' in paragraph 0 has Emphasis style (italic=True, style='Emphasis') — 0.34 pts
  Component 2: 'Canis lupus' in paragraph 2 has Emphasis style (italic=True, style='Emphasis') — 0.33 pts
  Component 3: 'Felis catus' in paragraph 4 has Emphasis style (italic=True, style='Emphasis') — 0.33 pts
  Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_054'


def has_emphasis_style(run):
    """
    Check if a run has the Emphasis character style applied.
    Emphasis style makes text italic. We check both the run style name
    and the italic property (either is acceptable evidence).
    """
    style_name = run.style.name if run.style else None
    is_emphasis_style = style_name == 'Emphasis'
    is_italic = run.font.italic is True
    return is_emphasis_style or is_italic


def find_run_with_text(para, target_text):
    """
    Find the run in a paragraph that contains the target_text exactly.
    Returns the run if found, None otherwise.
    """
    for run in para.runs:
        if run.text.strip() == target_text:
            return run
    # If no exact match, try partial match
    for run in para.runs:
        if target_text in run.text:
            return run
    return None


def verify_task(file_path):
    """
    Verify that the 'Emphasis' character style has been applied to all three
    species names in the biology notes document.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: verify document has at least 5 paragraphs
    if len(doc.paragraphs) < 5:
        print(f"CRITICAL: Document has only {len(doc.paragraphs)} paragraphs, expected at least 5")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: 'Homo sapiens' in paragraph 0 (1-indexed: paragraph 1) has Emphasis style (0.34 points)
    try:
        para0 = doc.paragraphs[0]
        homo_run = find_run_with_text(para0, 'Homo sapiens')
        if homo_run is not None:
            style_name = homo_run.style.name if homo_run.style else None
            italic_val = homo_run.font.italic
            if has_emphasis_style(homo_run):
                print(f"PASS: Component 1 — 'Homo sapiens' has Emphasis style "
                      f"(style='{style_name}', italic={italic_val}) (0.34 pts)")
                total_score += 0.34
            else:
                print(f"FAIL: Component 1 — 'Homo sapiens' does NOT have Emphasis style "
                      f"(style='{style_name}', italic={italic_val})")
        else:
            print("FAIL: Component 1 — 'Homo sapiens' run not found in paragraph 0")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 'Canis lupus' in paragraph 2 (1-indexed: paragraph 3) has Emphasis style (0.33 points)
    try:
        para2 = doc.paragraphs[2]
        canis_run = find_run_with_text(para2, 'Canis lupus')
        if canis_run is not None:
            style_name = canis_run.style.name if canis_run.style else None
            italic_val = canis_run.font.italic
            if has_emphasis_style(canis_run):
                print(f"PASS: Component 2 — 'Canis lupus' has Emphasis style "
                      f"(style='{style_name}', italic={italic_val}) (0.33 pts)")
                total_score += 0.33
            else:
                print(f"FAIL: Component 2 — 'Canis lupus' does NOT have Emphasis style "
                      f"(style='{style_name}', italic={italic_val})")
        else:
            print("FAIL: Component 2 — 'Canis lupus' run not found in paragraph 2")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: 'Felis catus' in paragraph 4 (1-indexed: paragraph 5) has Emphasis style (0.33 points)
    try:
        para4 = doc.paragraphs[4]
        felis_run = find_run_with_text(para4, 'Felis catus')
        if felis_run is not None:
            style_name = felis_run.style.name if felis_run.style else None
            italic_val = felis_run.font.italic
            if has_emphasis_style(felis_run):
                print(f"PASS: Component 3 — 'Felis catus' has Emphasis style "
                      f"(style='{style_name}', italic={italic_val}) (0.33 pts)")
                total_score += 0.33
            else:
                print(f"FAIL: Component 3 — 'Felis catus' does NOT have Emphasis style "
                      f"(style='{style_name}', italic={italic_val})")
        else:
            print("FAIL: Component 3 — 'Felis catus' run not found in paragraph 4")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path on the VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
