"""
Initial Setup: Biology notes document with animal classification paragraphs.
Species names are in regular 12pt Times New Roman (no Emphasis/italic style).
Task ID: writer_txtfmt_054
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_054'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DESKTOP_PATH = f'{WORKDIR}/Desktop/bio_notes.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_paragraph_with_species(doc, before_text, species_name, after_text, font_name='Times New Roman', font_size_pt=12):
    """Add a paragraph where the species name is in regular (non-italic) font."""
    para = doc.add_paragraph()
    # Text before species name
    if before_text:
        run_before = para.add_run(before_text)
        run_before.font.name = font_name
        run_before.font.size = Pt(font_size_pt)
        run_before.font.italic = False
    # Species name run - explicitly NOT italic, no Emphasis style
    run_species = para.add_run(species_name)
    run_species.font.name = font_name
    run_species.font.size = Pt(font_size_pt)
    run_species.font.italic = False
    # Text after species name
    if after_text:
        run_after = para.add_run(after_text)
        run_after.font.name = font_name
        run_after.font.size = Pt(font_size_pt)
        run_after.font.italic = False
    return para


def add_plain_paragraph(doc, text, font_name='Times New Roman', font_size_pt=12):
    """Add a simple paragraph with all regular text."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.italic = False
    return para


def create_initial():
    doc = Document()

    # Remove default empty paragraph from new Document
    # (python-docx adds one by default)
    # We'll just add our paragraphs and save

    # Clear existing paragraphs (the default empty one)
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # Paragraph 1: Contains 'Homo sapiens'
    add_paragraph_with_species(
        doc,
        before_text='Modern human beings, known scientifically as ',
        species_name='Homo sapiens',
        after_text=', belong to the family Hominidae and are the only extant members of the subtribe Hominina.',
        font_name='Times New Roman',
        font_size_pt=12,
    )

    # Paragraph 2: No species name
    add_plain_paragraph(
        doc,
        text='The classification of living organisms follows a hierarchical system developed by Carl Linnaeus in the 18th century. This system groups organisms based on shared characteristics and evolutionary relationships.',
        font_name='Times New Roman',
        font_size_pt=12,
    )

    # Paragraph 3: Contains 'Canis lupus'
    add_paragraph_with_species(
        doc,
        before_text='The gray wolf, or ',
        species_name='Canis lupus',
        after_text=', is a large canine native to Eurasia and North America. It is the ancestor of the domestic dog and plays a crucial role in maintaining ecological balance as an apex predator.',
        font_name='Times New Roman',
        font_size_pt=12,
    )

    # Paragraph 4: No species name
    add_plain_paragraph(
        doc,
        text='Binomial nomenclature, the formal system of naming species, uses a two-part name consisting of the genus and the specific epithet. This standardized naming convention ensures clarity in scientific communication across different languages.',
        font_name='Times New Roman',
        font_size_pt=12,
    )

    # Paragraph 5: Contains 'Felis catus'
    add_paragraph_with_species(
        doc,
        before_text='The domestic cat, ',
        species_name='Felis catus',
        after_text=', is a small, typically furry, carnivorous mammal. It is often called a house cat when kept as a pet or simply a cat. Domestication of cats is believed to have begun around 10,000 years ago in the Near East.',
        font_name='Times New Roman',
        font_size_pt=12,
    )

    # Save the main artifact
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Also place on Desktop as bio_notes.docx as referenced in the context
    import shutil
    os.makedirs(os.path.dirname(DESKTOP_PATH), exist_ok=True)
    shutil.copy(OUTPUT, DESKTOP_PATH)
    print(f'Copied to Desktop: {DESKTOP_PATH}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DESKTOP_PATH}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
