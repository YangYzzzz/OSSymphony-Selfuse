"""
Initial Setup: Create a Writer document with Default Page Style (portrait)
Task ID: writer_bs_059
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_059'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Default Page Style: Portrait, A4, standard margins ---
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # --- Title ---
    title = doc.add_heading('Quarterly Performance Review — Q1 2025', level=1)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run('Prepared by the Human Resources Department')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Introduction ---
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'This report summarizes the performance metrics across all departments '
        'for the first quarter of 2025. Overall, the company exceeded revenue '
        'targets by 8.3%, driven primarily by strong growth in the Asia-Pacific '
        'region and the successful launch of the Meridian product line.'
    )

    # --- Department Highlights ---
    doc.add_heading('Department Highlights', level=2)

    doc.add_heading('Engineering', level=3)
    doc.add_paragraph(
        'The engineering team completed 47 of 52 planned feature deliverables, '
        'achieving a 90.4% completion rate. Notable accomplishments include the '
        'migration of core services to the new cloud infrastructure, reducing '
        'average response times by 34%. Team lead Sarah Chen was recognized for '
        'her exceptional project management during the platform transition.'
    )

    doc.add_heading('Marketing', level=3)
    doc.add_paragraph(
        'Marketing campaigns generated 12,400 qualified leads, a 22% increase '
        'over Q4 2024. The digital advertising spend of $187,500 yielded a '
        'return on ad spend (ROAS) of 4.7x. Marcus Johnson led the rebranding '
        'initiative that contributed to a 15% improvement in brand recognition '
        'scores across key demographics.'
    )

    doc.add_heading('Sales', level=3)
    doc.add_paragraph(
        'Total revenue reached $3.42 million against a target of $3.16 million. '
        'The enterprise segment closed 8 new accounts valued at $1.85 million '
        'combined. Average deal size increased from $185,000 to $231,250. '
        'Regional manager Priya Patel secured the largest single contract in '
        'company history with Nexus Global Technologies.'
    )

    doc.add_heading('Customer Success', level=3)
    doc.add_paragraph(
        'Customer satisfaction (CSAT) scores averaged 4.6 out of 5.0 across '
        'all support channels. The team resolved 94% of critical tickets within '
        'the 4-hour SLA window. Net Promoter Score (NPS) improved from 62 to 71, '
        'reflecting the positive impact of the new onboarding process designed '
        'by team lead Kenji Watanabe.'
    )

    # --- Financial Table ---
    doc.add_heading('Financial Summary', level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ['Department', 'Budget ($)', 'Actual ($)', 'Variance (%)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Engineering', '1,250,000', '1,198,400', '-4.1'],
        ['Marketing', '420,000', '437,500', '+4.2'],
        ['Sales', '680,000', '652,300', '-4.1'],
        ['Customer Success', '310,000', '298,750', '-3.6'],
        ['Operations', '540,000', '561,200', '+3.9'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Conclusion ---
    doc.add_heading('Next Steps', level=2)
    doc.add_paragraph(
        'Based on the Q1 results, the leadership team recommends the following '
        'priorities for Q2 2025:'
    )
    doc.add_paragraph('Expand the Asia-Pacific sales team by 3 additional representatives', style='List Bullet')
    doc.add_paragraph('Allocate an additional $50,000 to digital marketing campaigns', style='List Bullet')
    doc.add_paragraph('Complete Phase 2 of the cloud infrastructure migration by June 30', style='List Bullet')
    doc.add_paragraph('Launch the revised customer onboarding workflow by May 15', style='List Bullet')
    doc.add_paragraph('Conduct mid-year performance reviews for all departments by July 1', style='List Bullet')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
