"""
Initial Setup: Release announcement template with placeholder text
Task ID: writer_tech_061
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_061'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # ----- Page setup -----
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ----- Title -----
    title = doc.add_heading('Product Release Announcement', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ----- Subtitle -----
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Nextera Platform — Quarterly Release Notification')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.italic = True

    # ----- Blank line -----
    doc.add_paragraph()

    # ----- Date line -----
    date_para = doc.add_paragraph()
    date_para.add_run('Date: ').bold = True
    date_para.add_run('April 2, 2026')

    # ----- Recipient greeting -----
    greeting = doc.add_paragraph()
    greeting.add_run('Dear ').font.size = Pt(11)
    run_placeholder = greeting.add_run('[Customer Name]')
    run_placeholder.font.size = Pt(11)
    greeting.add_run(',').font.size = Pt(11)

    # ----- Body paragraph 1 -----
    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after = Pt(6)
    body1.add_run(
        'We are pleased to inform you that Nextera Platform version '
    )
    run_ver = body1.add_run('[Version]')
    body1.add_run(
        ' is now available for deployment. This release includes significant '
        'performance improvements, new integrations with third-party analytics '
        'tools, and critical security patches that address vulnerabilities '
        'reported in our Q1 2026 security audit.'
    )

    # ----- Body paragraph 2 -----
    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after = Pt(6)
    body2.add_run(
        'The official release date is scheduled for '
    )
    run_date = body2.add_run('[Release Date]')
    body2.add_run(
        '. Your dedicated account team will coordinate the upgrade window '
        'to minimize downtime and ensure a smooth transition. We recommend '
        'reviewing the migration guide attached to this notification before '
        'the scheduled deployment.'
    )

    # ----- Key Highlights heading -----
    doc.add_heading('Key Highlights', level=1)

    # ----- Bullet list -----
    bullets = [
        'Enhanced API response times — up to 40% faster under peak load',
        'Native integration with Snowflake and Databricks data warehouses',
        'Upgraded TLS 1.3 support across all service endpoints',
        'New role-based access control (RBAC) for workspace administrators',
        'Automated compliance reporting for SOC 2 and ISO 27001',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # ----- Upgrade Instructions heading -----
    doc.add_heading('Upgrade Instructions', level=1)

    body3 = doc.add_paragraph()
    body3.paragraph_format.space_after = Pt(6)
    body3.add_run(
        'To begin the upgrade process, log in to the Nextera Admin Console '
        'and navigate to Settings > System Updates. Select version '
    )
    body3.add_run('[Version]')
    body3.add_run(
        ' from the available releases and follow the on-screen instructions. '
        'If you require assistance, please contact our support team at '
        'support@nextera.io or call +1 (800) 555-0199.'
    )

    # ----- Support & Contact heading -----
    doc.add_heading('Support & Contact', level=1)

    # ----- Table with support contacts -----
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['Department', 'Contact Person', 'Email']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    contacts = [
        ['Technical Support', 'Anika Patel', 'anika.patel@nextera.io'],
        ['Account Management', 'David Reyes', 'david.reyes@nextera.io'],
        ['Security Operations', 'Lin Wei', 'lin.wei@nextera.io'],
    ]
    for r, row_data in enumerate(contacts, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    doc.add_paragraph()

    # ----- Closing -----
    closing = doc.add_paragraph()
    closing.add_run(
        'Thank you for your continued partnership. We look forward to '
        'delivering even more value with this release.'
    )

    doc.add_paragraph()

    sign = doc.add_paragraph()
    sign.add_run('Best regards,').font.size = Pt(11)

    name_line = doc.add_paragraph()
    run_name = name_line.add_run('Jonathan Mitchell')
    run_name.bold = True
    run_name.font.size = Pt(11)

    title_line = doc.add_paragraph()
    run_title = title_line.add_run('VP of Product Engineering, Nextera Inc.')
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
