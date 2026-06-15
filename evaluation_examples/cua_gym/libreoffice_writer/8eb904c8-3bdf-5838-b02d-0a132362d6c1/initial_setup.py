"""
Initial Setup: Writer document with CodeBlock-styled paragraphs for macro task
Task ID: writer_tech_090
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_090'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_codeblock_style(doc):
    """Create a custom 'CodeBlock' paragraph style."""
    styles = doc.styles
    style = styles.add_style('CodeBlock', 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    style.base_style = styles['Normal']
    style.font.name = 'Courier New'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.left_indent = Inches(0.3)
    # Add shading to the style via XML
    pPr = style.element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F0F0'
    })
    pPr.append(shd)
    return style


def create_initial():
    doc = Document()

    # -- Title --
    title = doc.add_heading('Network Configuration Utility - Source Code Review', level=1)

    # -- Introduction paragraph --
    intro = doc.add_paragraph(
        'This document contains selected source code excerpts from the Network '
        'Configuration Utility project (v3.2.1). The code blocks below are marked '
        'with the CodeBlock style for easy identification during review sessions.'
    )

    doc.add_paragraph()  # spacer

    # -- Create CodeBlock style --
    cb_style = create_codeblock_style(doc)

    # -- Section 1 --
    doc.add_heading('1. Connection Handler Module', level=2)
    desc1 = doc.add_paragraph(
        'The following function manages TCP socket connections and implements '
        'retry logic with exponential backoff. Review for potential timeout issues.'
    )

    # Code Block 1 - Python connection handler (multi-line, no line numbers)
    code1 = doc.add_paragraph(style='CodeBlock')
    code1_text = (
        'import socket\n'
        'import time\n'
        'import logging\n'
        '\n'
        'def connect_with_retry(host, port, max_retries=5):\n'
        '    """Establish connection with exponential backoff."""\n'
        '    delay = 1.0\n'
        '    for attempt in range(max_retries):\n'
        '        try:\n'
        '            sock = socket.create_connection((host, port), timeout=10)\n'
        '            logging.info(f"Connected to {host}:{port}")\n'
        '            return sock\n'
        '        except socket.error as e:\n'
        '            logging.warning(f"Attempt {attempt+1} failed: {e}")\n'
        '            time.sleep(delay)\n'
        '            delay *= 2\n'
        '    raise ConnectionError(f"Failed after {max_retries} attempts")'
    )
    run1 = code1.add_run(code1_text)
    run1.font.name = 'Courier New'
    run1.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # -- Section 2 --
    doc.add_heading('2. Configuration Parser', level=2)
    desc2 = doc.add_paragraph(
        'This parser reads YAML-based network configuration files and validates '
        'the schema before applying changes. Pay attention to the error handling path.'
    )

    # Code Block 2 - Config parser
    code2 = doc.add_paragraph(style='CodeBlock')
    code2_text = (
        'class ConfigParser:\n'
        '    REQUIRED_KEYS = ["hostname", "port", "protocol"]\n'
        '\n'
        '    def __init__(self, config_path):\n'
        '        self.config_path = config_path\n'
        '        self.config = {}\n'
        '\n'
        '    def load(self):\n'
        '        with open(self.config_path, "r") as f:\n'
        '            self.config = yaml.safe_load(f)\n'
        '        self._validate()\n'
        '\n'
        '    def _validate(self):\n'
        '        missing = [k for k in self.REQUIRED_KEYS\n'
        '                   if k not in self.config]\n'
        '        if missing:\n'
        '            raise ValueError(f"Missing keys: {missing}")\n'
        '        if not 1 <= self.config["port"] <= 65535:\n'
        '            raise ValueError("Port out of range")'
    )
    run2 = code2.add_run(code2_text)
    run2.font.name = 'Courier New'
    run2.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # -- Section 3 --
    doc.add_heading('3. Packet Inspector', level=2)
    desc3 = doc.add_paragraph(
        'The packet inspector captures and analyzes network traffic on a specified '
        'interface. It filters packets by protocol type and logs anomalies.'
    )

    # Code Block 3 - Packet inspector
    code3 = doc.add_paragraph(style='CodeBlock')
    code3_text = (
        'def inspect_packets(interface, duration=60):\n'
        '    """Capture packets for analysis."""\n'
        '    cap = pcap.open_live(interface, 65536, True, 1000)\n'
        '    cap.setfilter("tcp or udp")\n'
        '    start_time = time.time()\n'
        '    packet_count = 0\n'
        '\n'
        '    while time.time() - start_time < duration:\n'
        '        header, data = cap.next()\n'
        '        if header is None:\n'
        '            continue\n'
        '        packet_count += 1\n'
        '        analyze_packet(data)\n'
        '\n'
        '    cap.close()\n'
        '    return packet_count'
    )
    run3 = code3.add_run(code3_text)
    run3.font.name = 'Courier New'
    run3.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # -- Closing section --
    doc.add_heading('Review Notes', level=2)
    closing = doc.add_paragraph(
        'Please review each code block above and provide feedback on error handling, '
        'performance characteristics, and adherence to the project coding standards. '
        'Use the comment feature to annotate specific lines of concern.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
