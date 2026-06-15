"""
Initial Setup: Create a Renovation Report document with heading and text, no tables.
Task ID: writer_rd_076
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_076'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    heading = doc.add_heading('Renovation Report', level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Introductory text ---
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(12)
    run = intro.add_run(
        'This report documents the comprehensive renovation of our corporate office space '
        'at 1450 Market Street, completed in Q1 2025. The project encompassed four major areas '
        'of the building, transforming outdated facilities into modern, collaborative workspaces '
        'designed to enhance productivity and employee well-being.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Project overview paragraph ---
    overview_heading = doc.add_heading('Project Overview', level=1)

    overview = doc.add_paragraph()
    overview.paragraph_format.space_after = Pt(10)
    run = overview.add_run(
        'The renovation was carried out by Henderson & Associates Architecture firm under '
        'the supervision of Project Manager Rachel Torres. With a total budget of $2.4 million, '
        'the project was completed on schedule over a 14-week period from November 2024 through '
        'February 2025. Each area was carefully redesigned to reflect our company\'s commitment '
        'to sustainability and modern design principles.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Scope paragraph ---
    scope = doc.add_paragraph()
    scope.paragraph_format.space_after = Pt(10)
    run = scope.add_run(
        'The scope of work included structural modifications, electrical upgrades, new HVAC '
        'installations, and complete interior finishing for the following areas: the main lobby '
        'and reception area, the executive conference room on the 3rd floor, the employee break '
        'room and kitchen, and the open office workspace on the 2nd floor.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Key highlights ---
    highlights_heading = doc.add_heading('Key Highlights', level=1)

    items = [
        'Energy-efficient LED lighting throughout all renovated spaces',
        'Acoustic panels installed in the conference room for improved sound quality',
        'Ergonomic furniture and standing desk options in the open office area',
        'Commercial-grade kitchen appliances in the redesigned break room',
        'Biophilic design elements including living walls in the lobby',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Calibri'

    # --- Note about visual comparison needed ---
    note_heading = doc.add_heading('Visual Comparison', level=1)

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    run = note.add_run(
        'A side-by-side comparison table showing before and after photographs of each '
        'renovated area is needed below to provide stakeholders with a clear visual '
        'representation of the transformation achieved in each space.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
