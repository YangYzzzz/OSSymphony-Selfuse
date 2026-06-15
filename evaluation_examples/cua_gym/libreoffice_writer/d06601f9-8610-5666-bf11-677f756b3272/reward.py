"""
Reward Script: Create a formatted comparison table in a Writer document
Task ID: writer_rd_076
Domain: libreoffice_writer
Scoring:
  Component 1: Table exists with correct dimensions (5 rows x 4 cols) — 0.20
  Component 2: Header row has merged 'Before' (cols 0-1) and 'After' (cols 2-3) — 0.20
  Component 3: Header formatting: red/green backgrounds, white bold text — 0.20
  Component 4: Data rows have correct area names and 'Image placeholder' text — 0.20
  Component 5: Alternating row backgrounds (white/light gray) — 0.20
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_076'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Component 1: Table has correct dimensions — 5 rows x 4 columns (0.20 points)
    try:
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        if num_rows == 5 and num_cols == 4:
            print(f"PASS: Component 1 — Table dimensions {num_rows}x{num_cols} correct (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Expected 5x4 table, found {num_rows}x{num_cols}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header row has merged cells — 'Before' spanning cols 0-1, 'After' spanning cols 2-3 (0.20 points)
    try:
        # Check 'Before' merge: cells [0,0] and [0,1] should have gridSpan=2
        before_tc = table.cell(0, 0)._tc
        before_tcPr = before_tc.find(qn('w:tcPr'))
        before_gs = None
        if before_tcPr is not None:
            gs_el = before_tcPr.find(qn('w:gridSpan'))
            if gs_el is not None:
                before_gs = gs_el.get(qn('w:val'))

        # Check 'After' merge: cells [0,2] and [0,3] should have gridSpan=2
        after_tc = table.cell(0, 2)._tc
        after_tcPr = after_tc.find(qn('w:tcPr'))
        after_gs = None
        if after_tcPr is not None:
            gs_el = after_tcPr.find(qn('w:gridSpan'))
            if gs_el is not None:
                after_gs = gs_el.get(qn('w:val'))

        before_text = table.cell(0, 0).text.strip().lower()
        after_text = table.cell(0, 2).text.strip().lower()

        if before_gs == '2' and after_gs == '2' and 'before' in before_text and 'after' in after_text:
            print(f"PASS: Component 2 — Header merged correctly: Before(span={before_gs}), After(span={after_gs}) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — Before span={before_gs} text='{before_text}', After span={after_gs} text='{after_text}'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header formatting — red bg for Before, green bg for After, white bold text (0.20 points)
    try:
        sub_score = 0.0

        # Check Before cell background (should be CC0000 or close to red)
        before_shd = None
        if before_tcPr is not None:
            shd_el = before_tcPr.find(qn('w:shd'))
            if shd_el is not None:
                before_shd = shd_el.get(qn('w:fill'))

        if before_shd and before_shd.upper() == 'CC0000':
            sub_score += 0.05
            print(f"  PASS: Before cell bg = #{before_shd}")
        else:
            print(f"  FAIL: Before cell bg expected #CC0000, found #{before_shd}")

        # Check After cell background (should be 006633 or close to green)
        after_shd = None
        if after_tcPr is not None:
            shd_el = after_tcPr.find(qn('w:shd'))
            if shd_el is not None:
                after_shd = shd_el.get(qn('w:fill'))

        if after_shd and after_shd.upper() == '006633':
            sub_score += 0.05
            print(f"  PASS: After cell bg = #{after_shd}")
        else:
            print(f"  FAIL: After cell bg expected #006633, found #{after_shd}")

        # Check white bold text in Before header
        before_has_white_bold = any(
            run.font.bold and run.font.color.rgb and str(run.font.color.rgb).upper() == 'FFFFFF'
            for para in table.cell(0, 0).paragraphs
            for run in para.runs
            if run.text.strip()
        )
        if before_has_white_bold:
            sub_score += 0.05
            print("  PASS: Before header has white bold text")
        else:
            print("  FAIL: Before header missing white bold text")

        # Check white bold text in After header
        after_has_white_bold = any(
            run.font.bold and run.font.color.rgb and str(run.font.color.rgb).upper() == 'FFFFFF'
            for para in table.cell(0, 2).paragraphs
            for run in para.runs
            if run.text.strip()
        )
        if after_has_white_bold:
            sub_score += 0.05
            print("  PASS: After header has white bold text")
        else:
            print("  FAIL: After header missing white bold text")

        if sub_score > 0:
            print(f"PASS: Component 3 — Header formatting ({sub_score} pts)")
            total_score += sub_score
        else:
            print("FAIL: Component 3 — No header formatting checks passed")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Data rows have correct area names and 'Image placeholder' text (0.20 points)
    try:
        expected_areas = ['lobby', 'conference room', 'break room', 'open office']
        areas_correct = 0

        for ri, area in enumerate(expected_areas, start=1):
            # Description columns: col 0 (Before desc) and col 2 (After desc)
            before_desc = table.cell(ri, 0).text.strip().lower()
            after_desc = table.cell(ri, 2).text.strip().lower()
            before_img = table.cell(ri, 1).text.strip().lower()
            after_img = table.cell(ri, 3).text.strip().lower()

            if (area in before_desc and area in after_desc and
                    'image placeholder' in before_img and 'image placeholder' in after_img):
                areas_correct += 1
            else:
                print(f"  FAIL: Row {ri} expected '{area}' + placeholders, found: "
                      f"desc=[{before_desc}, {after_desc}], img=[{before_img}, {after_img}]")

        area_score = 0.20 * (areas_correct / 4)
        if areas_correct == 4:
            print(f"PASS: Component 4 — All 4 data rows correct (0.20 pts)")
            total_score += area_score
        elif areas_correct > 0:
            print(f"PARTIAL: Component 4 — {areas_correct}/4 rows correct ({area_score:.2f} pts)")
            total_score += area_score
        else:
            print("FAIL: Component 4 — No data rows match expected content")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Alternating row backgrounds — odd data rows white, even data rows light gray #F2F2F2 (0.20 points)
    try:
        alt_correct = 0
        expected_fills = {
            1: 'FFFFFF',  # Row 1 (Lobby) — white
            2: 'F2F2F2',  # Row 2 (Conference Room) — light gray
            3: 'FFFFFF',  # Row 3 (Break Room) — white
            4: 'F2F2F2',  # Row 4 (Open Office) — light gray
        }

        for ri, expected_fill in expected_fills.items():
            cell_tc = table.cell(ri, 0)._tc
            cell_tcPr = cell_tc.find(qn('w:tcPr'))
            actual_fill = None
            if cell_tcPr is not None:
                shd_el = cell_tcPr.find(qn('w:shd'))
                if shd_el is not None:
                    actual_fill = shd_el.get(qn('w:fill'))

            if actual_fill and actual_fill.upper() == expected_fill:
                alt_correct += 1
            else:
                print(f"  FAIL: Row {ri} bg expected #{expected_fill}, found #{actual_fill}")

        alt_score = 0.20 * (alt_correct / 4)
        if alt_correct == 4:
            print(f"PASS: Component 5 — Alternating backgrounds correct (0.20 pts)")
            total_score += alt_score
        elif alt_correct > 0:
            print(f"PARTIAL: Component 5 — {alt_correct}/4 rows correct ({alt_score:.2f} pts)")
            total_score += alt_score
        else:
            print("FAIL: Component 5 — No alternating backgrounds detected")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Run verification
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
