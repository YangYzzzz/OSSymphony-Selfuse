"""
Initial Setup: Convert manually numbered definitions list to auto-numbered list
Task ID: writer_legal_028
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_028'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('Software License Agreement', level=0)

    # --- Preamble paragraph ---
    preamble = doc.add_paragraph(
        'This Software License Agreement ("Agreement") is entered into as of '
        'March 15, 2025, by and between Nextera Technologies, Inc., a Delaware '
        'corporation ("Licensor"), and Brightfield Solutions, LLC, a California '
        'limited liability company ("Licensee").'
    )

    # --- WHEREAS clauses ---
    doc.add_paragraph(
        'WHEREAS, Licensor owns certain proprietary software and related '
        'intellectual property; and'
    )
    doc.add_paragraph(
        'WHEREAS, Licensee desires to obtain a license to use such software '
        'subject to the terms and conditions set forth herein;'
    )
    doc.add_paragraph(
        'NOW, THEREFORE, in consideration of the mutual covenants contained '
        'herein and for other good and valuable consideration, the receipt and '
        'sufficiency of which are hereby acknowledged, the parties agree as follows:'
    )

    # --- Article I: Definitions ---
    doc.add_heading('Article I: Definitions', level=1)

    intro = doc.add_paragraph(
        'For purposes of this Agreement, the following terms shall have the '
        'meanings ascribed to them below:'
    )

    # 15 manually numbered definitions — plain Normal style, NO List Number
    definitions = [
        '1. "Affiliate" means any entity that directly or indirectly controls, '
        'is controlled by, or is under common control with a party, where '
        '"control" means ownership of more than fifty percent (50%) of the '
        'voting securities of such entity.',

        '2. "Agreement" means this Software License Agreement, including all '
        'exhibits, schedules, and amendments hereto.',

        '3. "Authorized Users" means the employees, contractors, and agents of '
        'Licensee who are authorized to access and use the Software pursuant '
        'to Section 3.1.',

        '4. "Confidential Information" means any non-public information disclosed '
        'by either party to the other party, whether orally, in writing, or by '
        'inspection, including but not limited to trade secrets, business plans, '
        'financial data, and technical specifications.',

        '5. "Documentation" means the user manuals, technical specifications, '
        'release notes, and other written materials provided by Licensor '
        'relating to the Software.',

        '6. "Effective Date" means the date first written above, being March 15, '
        '2025, or such later date as the parties may mutually agree in writing.',

        '7. "Intellectual Property Rights" means all patents, copyrights, '
        'trademarks, trade secrets, and other proprietary rights recognized '
        'under the laws of any jurisdiction worldwide.',

        '8. "License Fee" means the amounts payable by Licensee to Licensor '
        'as set forth in Exhibit A, currently totaling $245,000 per annum.',

        '9. "Licensed Territory" means the United States, Canada, the United '
        'Kingdom, and the European Economic Area, unless otherwise expanded '
        'by written amendment.',

        '10. "Maintenance Services" means the bug fixes, patches, updates, '
        'and technical support services provided by Licensor during the '
        'Maintenance Term as described in Section 7.2.',

        '11. "Permitted Purpose" means the internal business operations of '
        'Licensee, including data processing, reporting, and analytics, '
        'but excluding any resale, sublicensing, or external distribution.',

        '12. "Software" means the Nextera Analytics Platform, version 4.2, '
        'including all modules, components, libraries, and executable code '
        'delivered to Licensee.',

        '13. "Source Code" means the human-readable form of the Software, '
        'including all comments, annotations, and build scripts, which '
        'shall be held in escrow pursuant to Section 9.4.',

        '14. "Term" means the initial period of three (3) years commencing '
        'on the Effective Date, subject to renewal as set forth in Section '
        '12.1.',

        '15. "Update" means any modification, enhancement, or new release '
        'of the Software that Licensor makes generally available to its '
        'licensees during the Term.',
    ]

    for defn in definitions:
        doc.add_paragraph(defn)

    # --- Article II placeholder to add context ---
    doc.add_heading('Article II: Grant of License', level=1)
    doc.add_paragraph(
        'Subject to the terms and conditions of this Agreement, Licensor hereby '
        'grants to Licensee a non-exclusive, non-transferable license to use the '
        'Software within the Licensed Territory for the Permitted Purpose during '
        'the Term.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
