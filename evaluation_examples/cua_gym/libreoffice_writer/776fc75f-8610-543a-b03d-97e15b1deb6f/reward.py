"""
Reward Script: Set line spacing to 1.5 for body text in termination letter
Task ID: writer_hr_007
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): At least 20 Normal-style paragraphs have 1.5 line spacing
  Component 2 (0.4): Ratio of Normal-style paragraphs with 1.5 spacing >= 0.85
  Component 3 (0.2): Heading paragraphs remain at original (non-1.5) spacing
                      while body paragraphs are changed (compound check)
"""

import os

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_007'


def persist_app_state(domain: str):
    """Try to save any unsaved LibreOffice edits via Ctrl+S."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            import time
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Set line spacing to 1.5 lines for all body text paragraphs.
    Headings should remain unchanged (single spacing).
    """
    total_score = 0.0

    # Precondition: Load document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: Document must have paragraphs
    if len(doc.paragraphs) < 5:
        print(f"CRITICAL: Document has only {len(doc.paragraphs)} paragraphs, expected a full letter")
        print("REWARD: 0.0")
        return 0.0

    # Categorize paragraphs
    normal_paras = []
    heading_paras = []
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        if style_name.startswith("Heading"):
            heading_paras.append(para)
        else:
            normal_paras.append(para)

    print(f"INFO: Found {len(normal_paras)} Normal-style paragraphs, {len(heading_paras)} Heading paragraphs")

    # Count Normal paragraphs with 1.5 line spacing
    count_15 = 0
    count_checked = 0
    for para in normal_paras:
        pf = para.paragraph_format
        ls = pf.line_spacing
        count_checked += 1
        if ls is not None and abs(float(ls) - 1.5) < 0.01:
            count_15 += 1

    print(f"INFO: {count_15}/{count_checked} Normal paragraphs have 1.5 line spacing")

    # Component 1: At least 20 Normal-style paragraphs have 1.5 line spacing (0.4 points)
    # This checks absolute count - initial has 0, golden should have ~28
    try:
        if count_15 >= 20:
            print(f"PASS: Component 1 -- {count_15} Normal paragraphs have 1.5 spacing (>= 20 required) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 -- Only {count_15} Normal paragraphs have 1.5 spacing (>= 20 required)")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Ratio of Normal-style paragraphs with 1.5 spacing >= 0.85 (0.4 points)
    # This ensures comprehensive application, not just a few paragraphs
    try:
        if count_checked > 0:
            ratio = count_15 / count_checked
            if ratio >= 0.85:
                print(f"PASS: Component 2 -- {ratio:.2%} of Normal paragraphs have 1.5 spacing (>= 85% required) (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 2 -- Only {ratio:.2%} of Normal paragraphs have 1.5 spacing (>= 85% required)")
        else:
            print(f"FAIL: Component 2 -- No Normal paragraphs found")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Heading paragraphs remain at original spacing (not 1.5)
    #              AND at least 10 body paragraphs are at 1.5 (compound check) (0.2 points)
    # The compound condition ensures this only passes when the task change is present
    try:
        headings_with_15 = sum(
            1 for para in heading_paras
            if para.paragraph_format.line_spacing is not None
            and abs(float(para.paragraph_format.line_spacing) - 1.5) < 0.01
        )
        for para in heading_paras:
            ls = para.paragraph_format.line_spacing
            if ls is not None and abs(float(ls) - 1.5) < 0.01:
                print(f"  NOTE: Heading '{para.text[:40]}' has 1.5 spacing (should be unchanged)")

        if headings_with_15 == 0 and count_15 >= 10:
            print(f"PASS: Component 3 -- Headings unchanged AND {count_15} body paragraphs at 1.5 (0.2 pts)")
            total_score += 0.2
        elif headings_with_15 > 0:
            print(f"FAIL: Component 3 -- Some heading paragraphs were changed to 1.5 spacing")
        else:
            print(f"FAIL: Component 3 -- Not enough body paragraphs at 1.5 ({count_15} < 10)")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
