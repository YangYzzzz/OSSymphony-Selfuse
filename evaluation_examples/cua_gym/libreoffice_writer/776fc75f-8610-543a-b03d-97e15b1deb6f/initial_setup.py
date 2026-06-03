"""
Initial Setup: Termination letter with single line spacing
Task ID: writer_hr_007
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_007'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Company Heading ---
    heading = doc.add_heading('Meridian Technologies, Inc.', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Ensure single spacing on heading
    heading.paragraph_format.line_spacing = 1.0
    heading.paragraph_format.space_after = Pt(4)

    # Sub-heading
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.paragraph_format.line_spacing = 1.0
    sub.paragraph_format.space_after = Pt(12)
    run = sub.add_run('Human Resources Department')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.line_spacing = 1.0
    date_para.paragraph_format.space_after = Pt(6)
    run = date_para.add_run('March 28, 2026')
    run.font.size = Pt(11)

    # --- Recipient Address ---
    addr_lines = [
        'Mr. David R. Patterson',
        'Senior Software Developer',
        '4721 Elmwood Avenue, Apt 3B',
        'Portland, OR 97201',
    ]
    for line in addr_lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(11)

    # Add spacing after address block
    spacer = doc.add_paragraph()
    spacer.paragraph_format.line_spacing = 1.0
    spacer.paragraph_format.space_after = Pt(6)

    # --- Subject Line ---
    subject = doc.add_paragraph()
    subject.paragraph_format.line_spacing = 1.0
    subject.paragraph_format.space_after = Pt(6)
    run = subject.add_run('RE: Notice of Employment Termination')
    run.bold = True
    run.font.size = Pt(11)

    # --- Salutation ---
    sal = doc.add_paragraph()
    sal.paragraph_format.line_spacing = 1.0
    sal.paragraph_format.space_after = Pt(6)
    run = sal.add_run('Dear Mr. Patterson,')
    run.font.size = Pt(11)

    # --- Body Paragraphs (all single spacing) ---
    body_texts = [
        (
            'This letter serves as formal notification that your employment with '
            'Meridian Technologies, Inc. will be terminated effective April 11, 2026. '
            'This decision has been made after careful consideration and in accordance '
            'with the terms outlined in your employment agreement dated June 15, 2022.'
        ),
        (
            'As discussed during our meeting on March 25, 2026, with your direct '
            'supervisor, Ms. Angela Torres, and HR representative, Mr. Brian Whitfield, '
            'the company has undergone a significant restructuring of the Engineering '
            'division. Unfortunately, your position as Senior Software Developer in the '
            'Cloud Infrastructure team has been eliminated as part of this reorganization.'
        ),
        (
            'In recognition of your four years of service, Meridian Technologies is '
            'offering a severance package that includes the following provisions:'
        ),
        (
            '\u2022  Eight (8) weeks of base salary continuation at your current rate of '
            '$127,500 per annum, payable in regular bi-weekly installments beginning '
            'April 25, 2026.'
        ),
        (
            '\u2022  Continuation of health, dental, and vision insurance benefits through '
            'COBRA for a period of three (3) months, with the company covering the '
            'full premium cost through July 11, 2026.'
        ),
        (
            '\u2022  Accelerated vesting of 1,250 unvested stock options currently held under '
            'the 2023 Employee Equity Incentive Plan, subject to the terms of the plan '
            'agreement.'
        ),
        (
            '\u2022  Outplacement services through CareerBridge Partners for a period of '
            'six (6) months, including resume preparation, interview coaching, and '
            'job placement assistance.'
        ),
        (
            'Please note that acceptance of this severance package is contingent upon '
            'your execution of the enclosed Separation Agreement and General Release, '
            'which must be signed and returned to the Human Resources Department no '
            'later than April 18, 2026. You are encouraged to review this document '
            'with legal counsel of your choosing before signing.'
        ),
        (
            'Your final paycheck, including accrued and unused paid time off totaling '
            'twelve (12) days, will be issued on the next regular pay date following '
            'your last day of employment. Please ensure that all company property, '
            'including your laptop, security badge, parking pass, and any proprietary '
            'materials, is returned to the IT department by end of business on April 11, 2026.'
        ),
        (
            'We sincerely appreciate the contributions you have made to Meridian '
            'Technologies during your tenure. Your work on the Azure migration project '
            'and the internal developer tooling initiative were valuable to the organization. '
            'We wish you the very best in your future professional endeavors.'
        ),
    ]

    for text in body_texts:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(11)

    # --- Closing ---
    closing = doc.add_paragraph()
    closing.paragraph_format.line_spacing = 1.0
    closing.paragraph_format.space_before = Pt(12)
    closing.paragraph_format.space_after = Pt(0)
    run = closing.add_run('Sincerely,')
    run.font.size = Pt(11)

    # Signature lines
    sig_lines = [
        '',
        '',
        'Jennifer M. Caldwell',
        'Vice President, Human Resources',
        'Meridian Technologies, Inc.',
        'Direct: (503) 555-0184',
        'jennifer.caldwell@meridiantech.com',
    ]
    for line in sig_lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(11)

    # --- Enclosures note ---
    enc = doc.add_paragraph()
    enc.paragraph_format.line_spacing = 1.0
    enc.paragraph_format.space_before = Pt(12)
    enc.paragraph_format.space_after = Pt(0)
    run = enc.add_run('Enclosures: Separation Agreement and General Release')
    run.font.size = Pt(11)
    run.italic = True

    cc = doc.add_paragraph()
    cc.paragraph_format.line_spacing = 1.0
    cc.paragraph_format.space_after = Pt(0)
    run = cc.add_run('cc: Angela Torres, Engineering Director; Brian Whitfield, HR Manager; Legal Department')
    run.font.size = Pt(11)
    run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
