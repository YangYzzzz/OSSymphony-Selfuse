"""
Initial Setup: Insert page numbers in the header at the right side instead of the footer.
Task ID: writer_page_055
Domain: libreoffice_writer

Creates white_paper.docx on ~/Desktop/ with:
  - 7-page technology white paper content
  - A4 portrait, margins 2.54cm on all sides
  - No header
  - Footer enabled with centered page numbers
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'white_paper'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_field(paragraph, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER):
    """Add a PAGE field code to a paragraph (for footer page numbers)."""
    paragraph.paragraph_format.alignment = alignment
    paragraph.clear()

    run_begin = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run_begin._element.append(fld_char_begin)

    run_instr = paragraph.add_run()
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE \u005cArabic '
    run_instr._element.append(instr_text)

    run_end = paragraph.add_run()
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run_end._element.append(fld_char_end)


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)
    doc = Document()

    # ── Page setup: A4 portrait, 2.54 cm margins ──────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

    # ── No header: make sure header is empty / not linked ─────────────────────
    header = section.header
    header.is_linked_to_previous = False
    # Clear any default content in header
    for para in header.paragraphs:
        for run in para.runs:
            run.text = ''
        para.clear()

    # ── Footer: centered page numbers ─────────────────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    add_page_number_field(fp, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)

    # ── Document content: 7-page technology white paper ───────────────────────

    # Title page
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    title_run = title_para.add_run('Transforming Enterprise Operations Through\nArtificial Intelligence and Cloud-Native Architecture')
    title_run.bold = True
    title_run.font.size = Pt(24)

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_para.paragraph_format.space_before = Pt(24)
    subtitle_run = subtitle_para.add_run('A Technology White Paper')
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True

    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.space_before = Pt(48)
    author_run = author_para.add_run(
        'Dr. Megan Harrington, Principal Architect\n'
        'James Okafor, Senior Cloud Engineer\n'
        'Priya Nambiar, Head of Data Science\n\n'
        'TechVision Research Institute\nMarch 2025'
    )
    author_run.font.size = Pt(12)

    doc.add_page_break()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This white paper examines the convergence of artificial intelligence (AI) and '
        'cloud-native architecture as a transformative force for enterprise operations. '
        'Drawing on case studies from financial services, healthcare, and manufacturing, '
        'we demonstrate that organizations adopting integrated AI-cloud strategies achieve '
        'an average of 34% reduction in operational costs and 2.7x improvement in '
        'time-to-market for new product capabilities.'
    )
    doc.add_paragraph(
        'Key findings include: (1) microservices-based AI pipelines reduce model deployment '
        'latency by 68% compared to monolithic deployments; (2) federated learning frameworks '
        'enable privacy-preserving model training across distributed data silos; and '
        '(3) FinOps-driven cloud governance reduces wasted AI compute spend by up to 41%.'
    )

    doc.add_page_break()

    # Section 1
    doc.add_heading('1. Introduction: The Convergence Imperative', level=1)
    doc.add_paragraph(
        'The past decade has witnessed unprecedented growth in both AI capabilities and '
        'cloud computing infrastructure. What began as parallel technological trajectories '
        'have converged into a mutually reinforcing ecosystem that is redefining competitive '
        'advantage across industries. Enterprise leaders face a strategic inflection point: '
        'adapt to this new paradigm or risk obsolescence.'
    )
    doc.add_heading('1.1 Market Context', level=2)
    doc.add_paragraph(
        'Global enterprise AI spending reached $154.4 billion in 2024, with cloud-based '
        'AI services accounting for 71% of total investment (IDC, 2025). The compound '
        'annual growth rate of 28.3% is projected to sustain through 2030, driven by '
        'advances in large language models, computer vision, and autonomous decision systems.'
    )
    doc.add_paragraph(
        'Simultaneously, cloud infrastructure has matured beyond simple lift-and-shift '
        'migrations to encompass sophisticated orchestration platforms including Kubernetes, '
        'service meshes, and serverless compute. These capabilities provide the elastic '
        'scalability essential for training and serving large-scale AI models.'
    )

    doc.add_heading('1.2 The Integration Challenge', level=2)
    doc.add_paragraph(
        'Despite the compelling business case, 62% of enterprises report significant '
        'challenges integrating AI workloads with existing cloud infrastructure (Gartner, 2025). '
        'Common barriers include data gravity constraints, MLOps toolchain fragmentation, '
        'regulatory compliance requirements, and the organizational skills gap between '
        'data science and platform engineering teams.'
    )

    doc.add_page_break()

    # Section 2
    doc.add_heading('2. Cloud-Native AI Architecture Patterns', level=1)
    doc.add_paragraph(
        'Successful enterprise AI deployments share common architectural characteristics '
        'that leverage cloud-native principles to deliver reliability, scalability, and '
        'operational efficiency. This section details the core patterns observed in '
        'high-performing organizations.'
    )
    doc.add_heading('2.1 Microservices-Based ML Pipelines', level=2)
    doc.add_paragraph(
        'Traditional monolithic ML pipelines suffer from tight coupling between data '
        'ingestion, feature engineering, model training, and inference components. '
        'Decomposing these into independent microservices enables parallel development, '
        'independent scaling, and fault isolation. Leading practitioners decompose '
        'pipelines into feature stores, training orchestrators, model registries, '
        'and inference serving layers.'
    )
    doc.add_heading('2.2 Event-Driven Model Serving', level=2)
    doc.add_paragraph(
        'Synchronous REST-based model serving creates latency bottlenecks under high '
        'concurrency. Event-driven architectures using Apache Kafka or AWS Kinesis '
        'decouple prediction requests from model execution, enabling horizontal scaling '
        'without service degradation. Asynchronous serving patterns achieve 4.2x higher '
        'throughput than synchronous equivalents under peak load conditions.'
    )
    doc.add_heading('2.3 Federated Learning Infrastructure', level=2)
    doc.add_paragraph(
        'Federated learning addresses data privacy requirements by training models across '
        'distributed data sources without centralizing sensitive information. Cloud-native '
        'federated infrastructure requires secure aggregation protocols, differential '
        'privacy mechanisms, and edge-compatible model architectures. Organizations in '
        'healthcare and finance report 89% faster regulatory approval for federated '
        'models compared to centralized alternatives.'
    )

    doc.add_page_break()

    # Section 3
    doc.add_heading('3. Implementation Case Studies', level=1)
    doc.add_heading('3.1 GlobalBank Financial Services', level=2)
    doc.add_paragraph(
        'GlobalBank deployed a cloud-native fraud detection system processing 2.4 million '
        'transactions per second across 47 markets. The system employs ensemble models '
        'combining gradient boosting, transformer-based sequence analysis, and graph '
        'neural networks. Key architectural decisions included multi-region active-active '
        'deployment, sub-10ms inference SLOs, and automated model retraining triggered '
        'by data drift detection.'
    )
    doc.add_paragraph(
        'Results: 94.7% fraud detection accuracy (up from 87.2%), $340 million annual '
        'fraud loss reduction, 23% decrease in false positive chargebacks improving '
        'customer satisfaction scores from 71 to 84 NPS points.'
    )
    doc.add_heading('3.2 MediCare Health Network', level=2)
    doc.add_paragraph(
        'MediCare implemented a federated AI platform connecting 312 hospitals and '
        '1,847 clinics while maintaining HIPAA compliance and data sovereignty requirements. '
        'The platform enables collaborative model training for radiology image analysis, '
        'patient readmission prediction, and drug interaction detection without sharing '
        'patient records across institutional boundaries.'
    )
    doc.add_paragraph(
        'Outcomes include: 31% improvement in early-stage cancer detection rates, '
        '18% reduction in unplanned readmissions, and $127 million annual savings in '
        'preventable adverse drug events across the network.'
    )

    doc.add_page_break()

    # Section 4
    doc.add_heading('4. FinOps for AI Workloads', level=1)
    doc.add_paragraph(
        'AI and ML workloads exhibit distinct cost characteristics compared to traditional '
        'enterprise applications. Training workloads are bursty and compute-intensive, '
        'while inference serving demands consistent low-latency capacity. Effective FinOps '
        'practices must account for both dimensions.'
    )
    doc.add_heading('4.1 Spot Instance Strategies for Training', level=2)
    doc.add_paragraph(
        'GPU-accelerated training on spot or preemptible instances reduces compute costs '
        'by 60-80% compared to on-demand pricing. Checkpoint-based fault tolerance enables '
        'training resumption after instance interruption with minimal work loss. Advanced '
        'strategies include heterogeneous spot fleets, multi-zone capacity diversification, '
        'and predictive interruption avoidance using historical spot market signals.'
    )
    doc.add_heading('4.2 Inference Cost Optimization', level=2)
    doc.add_paragraph(
        'Model compression techniques including quantization, pruning, and knowledge '
        'distillation reduce inference compute requirements by 40-75% with minimal '
        'accuracy degradation. Hardware-software co-optimization using custom silicon '
        '(AWS Inferentia, Google TPU) delivers additional 3-5x cost efficiency improvements '
        'for high-volume inference workloads.'
    )

    doc.add_page_break()

    # Section 5 + Conclusion
    doc.add_heading('5. Security and Governance Framework', level=1)
    doc.add_paragraph(
        'Enterprise AI deployments require robust security controls spanning model '
        'integrity, data lineage, access governance, and adversarial robustness. '
        'Cloud-native security tooling provides foundational controls, but AI-specific '
        'requirements demand additional layers of protection.'
    )
    doc.add_heading('5.1 Model Security', level=2)
    doc.add_paragraph(
        'Model supply chain attacks represent an emerging threat vector. Organizations '
        'must implement model signing, provenance tracking, and behavioral integrity '
        'monitoring. Cryptographic attestation of model artifacts from training through '
        'deployment creates an auditable chain of custody essential for regulated industries.'
    )
    doc.add_heading('5.2 AI Governance and Responsible Use', level=2)
    doc.add_paragraph(
        'Automated governance pipelines enforce fairness constraints, explainability '
        'requirements, and regulatory compliance checks at each stage of the ML lifecycle. '
        'Model cards and datasheets provide standardized documentation for AI systems, '
        'enabling informed deployment decisions and stakeholder accountability.'
    )

    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The convergence of AI and cloud-native architecture represents the defining '
        'technology paradigm of the current enterprise era. Organizations that successfully '
        'navigate this convergence — building integrated platforms, developing cross-functional '
        'teams, and establishing adaptive governance frameworks — will achieve durable '
        'competitive advantages that compound over time.'
    )
    doc.add_paragraph(
        'The path forward requires deliberate investment in architectural foundations, '
        'organizational capability, and responsible AI practices. The case studies and '
        'patterns presented in this white paper provide a practical roadmap for enterprise '
        'technology leaders beginning or accelerating this transformative journey.'
    )

    doc.add_heading('References', level=1)
    refs = [
        'IDC (2025). Worldwide Artificial Intelligence Spending Guide. International Data Corporation.',
        'Gartner (2025). Magic Quadrant for Cloud AI Developer Services. Gartner Research.',
        'Dean, J., Ghemawat, S. (2024). Large-Scale Distributed Deep Networks. Proceedings of NIPS.',
        'McMahan, B. et al. (2023). Communication-Efficient Learning of Deep Networks from Decentralized Data. AISTATS.',
        'Sculley, D. et al. (2024). Hidden Technical Debt in Machine Learning Systems. NIPS.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style='List Number')
        p.paragraph_format.space_after = Pt(4)

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # ── GUI-ready startup ──────────────────────────────────────────────────────
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
