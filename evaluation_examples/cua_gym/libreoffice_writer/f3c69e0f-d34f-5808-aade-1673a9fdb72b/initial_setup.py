"""
Initial Setup: Governance amendment document with no comments
Task ID: writer_struct_064
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_064'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/governance_amendment.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set document margins
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_heading('GOVERNANCE AMENDMENT AND OVERSIGHT POLICY', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(12)

    subtitle = doc.add_paragraph('Resolution No. 2025-04 | Effective Date: April 1, 2025')
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(18)

    # Section 1 heading
    doc.add_heading('Section 1: Preamble and Authority', level=1)

    # Paragraph 1
    p1 = doc.add_paragraph(
        'Pursuant to the authority vested in the Board of Directors under the organization\'s '
        'Articles of Incorporation and applicable state law, this amendment is hereby adopted to '
        'establish updated governance procedures for the oversight committee. This resolution '
        'supersedes all prior resolutions in conflict herewith and shall be effective upon the '
        'date set forth above.'
    )
    p1.paragraph_format.space_after = Pt(8)

    # Paragraph 2 — CRITICAL: contains 'restructuring'
    p2 = doc.add_paragraph(
        'The restructuring of the oversight committee shall commence upon approval. '
        'All affected personnel are to be notified within ten (10) business days of the '
        'adoption of this resolution. Department heads are required to submit transition '
        'plans to the newly appointed committee chair no later than thirty (30) days '
        'following commencement.'
    )
    p2.paragraph_format.space_after = Pt(8)

    # Paragraph 3
    p3 = doc.add_paragraph(
        'Membership of the oversight committee shall consist of no fewer than five (5) '
        'and no more than nine (9) voting members, each appointed by majority vote of the '
        'full Board. Members shall serve staggered terms of three (3) years and may be '
        'reappointed for one additional consecutive term.'
    )
    p3.paragraph_format.space_after = Pt(8)

    # Section 2 heading
    doc.add_heading('Section 2: Duties and Responsibilities', level=1)

    # Paragraph 4
    p4 = doc.add_paragraph(
        'The oversight committee shall be responsible for monitoring compliance with all '
        'applicable laws, regulations, and internal policies. The committee shall convene '
        'quarterly and submit written reports to the Board within fifteen (15) days following '
        'each meeting. Ad hoc meetings may be called by the Committee Chair or by written '
        'request of at least three (3) members.'
    )
    p4.paragraph_format.space_after = Pt(8)

    # Paragraph 5
    p5 = doc.add_paragraph(
        'The committee shall have authority to engage independent counsel, auditors, and '
        'other subject matter experts as deemed necessary and appropriate. Reasonable and '
        'customary fees for such engagements shall be borne by the organization\'s general '
        'operating budget, subject to the approval of the Chief Financial Officer for '
        'expenditures exceeding twenty-five thousand dollars ($25,000).'
    )
    p5.paragraph_format.space_after = Pt(8)

    # Paragraph 6
    p6 = doc.add_paragraph(
        'Annual performance evaluations of the committee shall be conducted by the '
        'Governance and Nominating Committee and presented to the full Board in conjunction '
        'with the annual meeting. Criteria for evaluation shall include attendance records, '
        'quality of reports submitted, timeliness of action items, and overall effectiveness '
        'of oversight activities undertaken during the evaluation period.'
    )
    p6.paragraph_format.space_after = Pt(8)

    # Page break before Section 3
    doc.add_page_break()

    # Section 3 heading
    doc.add_heading('Section 3: Conflict of Interest and Recusal Policy', level=1)

    # Paragraph 7
    p7 = doc.add_paragraph(
        'All members of the oversight committee are required to disclose any actual or '
        'potential conflicts of interest in accordance with the organization\'s Conflict of '
        'Interest Policy. Disclosure must be made in writing to the Committee Chair and to '
        'the General Counsel no later than five (5) business days after the member becomes '
        'aware of the potential conflict.'
    )
    p7.paragraph_format.space_after = Pt(8)

    # Paragraph 8
    p8 = doc.add_paragraph(
        'A member who has disclosed a conflict of interest shall recuse themselves from '
        'all deliberations and votes on the matter to which the conflict relates. '
        'Documentation of such recusals shall be maintained in the committee\'s official '
        'minutes and retained for a period of not less than seven (7) years.'
    )
    p8.paragraph_format.space_after = Pt(8)

    # Paragraph 9
    p9 = doc.add_paragraph(
        'Failure to disclose a conflict of interest may result in disciplinary action, '
        'including removal from the committee, as determined by a two-thirds (2/3) vote of '
        'the full Board. The affected member shall have the right to present a written '
        'response prior to any such vote being conducted.'
    )
    p9.paragraph_format.space_after = Pt(8)

    # Section 4 heading
    doc.add_heading('Section 4: Amendment and Review Procedures', level=1)

    # Paragraph 10
    p10 = doc.add_paragraph(
        'This governance policy shall be reviewed no less frequently than every two (2) years '
        'by the Governance and Nominating Committee. Proposed amendments shall be submitted '
        'in writing to the Board Secretary no later than thirty (30) days prior to the Board '
        'meeting at which the amendment is to be considered.'
    )
    p10.paragraph_format.space_after = Pt(8)

    # Paragraph 11
    p11 = doc.add_paragraph(
        'Amendments to this policy require the affirmative vote of two-thirds (2/3) of the '
        'directors then in office. An emergency amendment may be adopted by unanimous written '
        'consent of all directors if circumstances require immediate action prior to the next '
        'regularly scheduled Board meeting.'
    )
    p11.paragraph_format.space_after = Pt(8)

    # Page break before final section
    doc.add_page_break()

    # Section 5 heading
    doc.add_heading('Section 5: Signatures and Attestations', level=1)

    # Paragraph 12
    p12 = doc.add_paragraph(
        'This resolution was duly adopted at a meeting of the Board of Directors at which a '
        'quorum was present and acting throughout. The undersigned officers of the organization '
        'hereby certify that the foregoing is a true and complete copy of the resolution duly '
        'adopted by the Board of Directors.'
    )
    p12.paragraph_format.space_after = Pt(24)

    # Signature block
    sig_lines = [
        '___________________________          ___________________________',
        'Chair, Board of Directors            Secretary, Board of Directors',
        '',
        '___________________________          ___________________________',
        'Date                                 Date',
        '',
        '___________________________',
        'General Counsel',
        '',
        '___________________________',
        'Date',
    ]
    for line in sig_lines:
        sp = doc.add_paragraph(line)
        sp.paragraph_format.space_after = Pt(4)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
