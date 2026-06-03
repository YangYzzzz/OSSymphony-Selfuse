"""
Initial Setup: Technical specification document with introduction paragraph.
Task ID: osworld_writer_image_insertion_004
Domain: libreoffice_writer
Description: Creates a technical specification .docx with intro paragraph and a
             diagram.png file in Documents folder. No image inserted in the doc yet.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_image_insertion_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
IMAGE_PATH = f'{WORKDIR}/Documents/diagram.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_diagram_png():
    """Create a simple diagram PNG in the Documents folder."""
    from PIL import Image, ImageDraw, ImageFont

    os.makedirs(f'{WORKDIR}/Documents', exist_ok=True)
    # Create a 600x400 diagram image
    img = Image.new('RGB', (600, 400), color=(240, 248, 255))
    draw = ImageDraw.Draw(img)

    # Draw a simple system architecture diagram
    # Main box
    draw.rectangle([50, 50, 550, 350], outline=(70, 130, 180), width=3)
    draw.rectangle([50, 50, 550, 100], fill=(70, 130, 180), outline=(70, 130, 180), width=2)

    # Title bar
    draw.text((300, 75), "System Architecture", fill=(255, 255, 255), anchor='mm')

    # Component boxes
    draw.rectangle([80, 130, 230, 200], fill=(173, 216, 230), outline=(70, 130, 180), width=2)
    draw.text((155, 165), "Frontend", fill=(0, 0, 0), anchor='mm')

    draw.rectangle([260, 130, 340, 200], fill=(144, 238, 144), outline=(34, 139, 34), width=2)
    draw.text((300, 165), "API", fill=(0, 0, 0), anchor='mm')

    draw.rectangle([370, 130, 520, 200], fill=(255, 218, 185), outline=(210, 105, 30), width=2)
    draw.text((445, 165), "Database", fill=(0, 0, 0), anchor='mm')

    # Arrows
    draw.line([230, 165, 260, 165], fill=(0, 0, 0), width=2)
    draw.line([340, 165, 370, 165], fill=(0, 0, 0), width=2)

    # Bottom layer
    draw.rectangle([150, 260, 450, 320], fill=(221, 160, 221), outline=(148, 0, 211), width=2)
    draw.text((300, 290), "Infrastructure Layer", fill=(0, 0, 0), anchor='mm')

    # Vertical arrows
    draw.line([300, 200, 300, 260], fill=(0, 0, 0), width=2)

    img.save(IMAGE_PATH)
    print(f'Diagram image created: {IMAGE_PATH}')


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('DataSync Pro — Technical Specification', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Introduction section ---
    heading_intro = doc.add_heading('1. Introduction', level=1)

    intro_para = doc.add_paragraph(
        'DataSync Pro is a distributed data synchronization platform designed to enable '
        'real-time data consistency across heterogeneous enterprise systems. This document '
        'provides a comprehensive technical specification covering system architecture, '
        'component interactions, deployment considerations, and integration guidelines. '
        'The platform supports high-throughput workloads with sub-second latency guarantees '
        'and is built upon a microservices architecture that ensures horizontal scalability '
        'and fault tolerance across all critical data pathways.'
    )
    intro_para.paragraph_format.space_after = Pt(6)

    # NOTE: No image is inserted here — the agent must insert diagram.png

    # --- Architecture Overview section ---
    doc.add_heading('2. Architecture Overview', level=1)
    arch_para = doc.add_paragraph(
        'The system is composed of three primary layers: the presentation layer, the '
        'business logic layer, and the data persistence layer. Each layer is independently '
        'scalable and communicates via well-defined RESTful APIs and message queues. '
        'The architecture follows the CQRS pattern to separate read and write operations, '
        'optimizing performance for both query-intensive and write-intensive workloads.'
    )
    arch_para.paragraph_format.space_after = Pt(6)

    # --- Component Details section ---
    doc.add_heading('3. Component Details', level=1)

    doc.add_heading('3.1 Frontend Service', level=2)
    doc.add_paragraph(
        'The frontend service provides the user-facing interface and implements a '
        'reactive, component-based architecture using modern web technologies. It '
        'communicates exclusively through the API gateway and maintains a local state '
        'cache to minimize redundant network requests.'
    )

    doc.add_heading('3.2 API Gateway', level=2)
    doc.add_paragraph(
        'The API gateway acts as the single entry point for all client requests. '
        'It handles authentication, rate limiting, request routing, and response '
        'aggregation. Circuit breakers are implemented at each upstream service '
        'connection to prevent cascade failures.'
    )

    doc.add_heading('3.3 Database Layer', level=2)
    doc.add_paragraph(
        'The database layer employs a polyglot persistence strategy, utilizing a '
        'relational database for transactional data, a document store for unstructured '
        'content, and an in-memory cache for frequently accessed records. Data '
        'replication ensures cross-region availability with automatic failover.'
    )

    # --- Performance Requirements ---
    doc.add_heading('4. Performance Requirements', level=1)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Metric', 'Target', 'Measurement Method']
    for col, h in enumerate(headers):
        cell = table.cell(0, col)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    perf_data = [
        ['API Response Time (P99)', '< 200 ms', 'Prometheus histogram'],
        ['Throughput', '10,000 req/sec', 'Load test (k6)'],
        ['Data Sync Latency', '< 500 ms', 'End-to-end trace'],
        ['System Availability', '99.95%', 'Uptime monitoring'],
    ]
    for r, row_data in enumerate(perf_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()

    # --- Deployment ---
    doc.add_heading('5. Deployment', level=1)
    doc.add_paragraph(
        'All services are containerized using Docker and orchestrated via Kubernetes. '
        'Infrastructure-as-code definitions are maintained in Terraform. Blue-green '
        'deployments are used to ensure zero-downtime releases, with automated rollback '
        'triggered when error-rate thresholds are exceeded during the promotion window.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


# Create diagram.png first
create_diagram_png()
# Then create the document
create_initial()
