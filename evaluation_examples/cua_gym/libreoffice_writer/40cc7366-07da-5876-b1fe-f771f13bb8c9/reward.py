"""
Reward Script: Insert image 'diagram.png' centered below introduction paragraph,
               set to 12 cm wide and 8 cm tall, add caption 'Figure 1: System Architecture',
               set wrapping to 'None' with center alignment.
Task ID: osworld_writer_image_insertion_004
Domain: libreoffice_writer
Scoring:
  Component 1: Image is inserted in the document (0.30 pts)
  Component 2: Image dimensions are exactly 12 cm wide x 8 cm tall (0.30 pts)
  Component 3: Caption 'Figure 1: System Architecture' is present and centered (0.20 pts)
  Component 4: Image paragraph has center alignment (0.10 pts)
  Component 5: Image uses inline (None) wrapping — wp:inline element used (0.10 pts)
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_image_insertion_004'

# EMU values for 12 cm and 8 cm (1 cm = 360000 EMU)
EXPECTED_WIDTH_EMU = 4320000   # 12 cm
EXPECTED_HEIGHT_EMU = 2880000  # 8 cm
EMU_TOLERANCE = 36000          # +/- 0.1 cm tolerance

CAPTION_TEXT = 'Figure 1: System Architecture'


def persist_app_state():
    """Best-effort save of any open LibreOffice Writer window."""
    try:
        import time
        os.environ["DISPLAY"] = ":0"
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.5)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def find_image_paragraphs(doc):
    """Return list of (index, para) for paragraphs containing inline or anchor drawings."""
    result = []
    for i, para in enumerate(doc.paragraphs):
        xml = para._element.xml
        if 'wp:inline' in xml or 'wp:anchor' in xml:
            result.append((i, para))
    return result


def get_image_extent(para):
    """Extract (cx, cy) EMU from the first drawing extent found in a paragraph."""
    from lxml import etree
    NS_WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    extents = para._element.findall(f'.//{{{NS_WP}}}extent')
    for extent in extents:
        cx_str = extent.get('cx')
        cy_str = extent.get('cy')
        if cx_str is not None and cy_str is not None:
            return int(cx_str), int(cy_str)
    return None, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: Image is inserted in the document (0.30 pts)
    # Verify both: (a) image relationship in doc.part.rels, AND
    #              (b) a paragraph contains a drawing element.
    # This fails on initial_env (no image) and passes on golden_env.
    # -----------------------------------------------------------------------
    try:
        image_rels = [
            rel for rel in doc.part.rels.values()
            if "image" in rel.reltype
        ]
        image_paras = find_image_paragraphs(doc)
        if len(image_rels) > 0 and len(image_paras) > 0:
            print(f"PASS: Component 1 — Image present: {len(image_rels)} rel(s), "
                  f"drawing in para index(es) {[idx for idx, _ in image_paras]} (0.30 pts)")
            total_score += 0.30
        else:
            if len(image_rels) == 0:
                print("FAIL: Component 1 — No image relationships found in document")
            else:
                print("FAIL: Component 1 — Image rel exists but no paragraph has a drawing element")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: Image dimensions exactly 12 cm wide x 8 cm tall (0.30 pts)
    # Reads wp:extent cx/cy attributes from the inline drawing element.
    # 12 cm = 4320000 EMU; 8 cm = 2880000 EMU.
    # -----------------------------------------------------------------------
    try:
        cx, cy = None, None
        for _idx, para in find_image_paragraphs(doc):
            cx, cy = get_image_extent(para)
            if cx is not None:
                break
        if cx is not None and cy is not None:
            cx_ok = abs(cx - EXPECTED_WIDTH_EMU) <= EMU_TOLERANCE
            cy_ok = abs(cy - EXPECTED_HEIGHT_EMU) <= EMU_TOLERANCE
            if cx_ok and cy_ok:
                print(f"PASS: Component 2 — Dimensions: cx={cx} (~{cx/360000:.2f} cm), "
                      f"cy={cy} (~{cy/360000:.2f} cm) [expected 12.00x8.00 cm] (0.30 pts)")
                total_score += 0.30
            else:
                print(f"FAIL: Component 2 — Wrong dimensions: cx={cx} (~{cx/360000:.2f} cm), "
                      f"cy={cy} (~{cy/360000:.2f} cm); expected 12.00 cm x 8.00 cm")
        else:
            print("FAIL: Component 2 — No drawing extent attributes found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Caption 'Figure 1: System Architecture' present and
    #              center-aligned (0.20 pts)
    # -----------------------------------------------------------------------
    try:
        caption_para = None
        for para in doc.paragraphs:
            if CAPTION_TEXT in para.text:
                caption_para = para
                break
        if caption_para is not None:
            cap_align = caption_para.paragraph_format.alignment
            if cap_align == WD_PARAGRAPH_ALIGNMENT.CENTER:
                print(f"PASS: Component 3 — Caption '{CAPTION_TEXT}' found and center-aligned (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 — Caption found but alignment={cap_align}, expected CENTER")
        else:
            candidates = [p.text for p in doc.paragraphs if 'Figure' in p.text or 'Architecture' in p.text]
            print(f"FAIL: Component 3 — Caption '{CAPTION_TEXT}' not found; candidates: {candidates}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: Image paragraph has center alignment (0.10 pts)
    # -----------------------------------------------------------------------
    try:
        img_paras = find_image_paragraphs(doc)
        if len(img_paras) > 0:
            _idx, first_img_para = img_paras[0]
            img_align = first_img_para.paragraph_format.alignment
            if img_align == WD_PARAGRAPH_ALIGNMENT.CENTER:
                print("PASS: Component 4 — Image paragraph is center-aligned (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 4 — Image paragraph alignment={img_align}, expected CENTER")
        else:
            print("FAIL: Component 4 — No image paragraph found to check alignment")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -----------------------------------------------------------------------
    # Component 5: Image uses inline (None) wrapping — wp:inline (0.10 pts)
    # wp:inline = in-line (no text wrap); wp:anchor = floating/wrap mode.
    # Count occurrences of each tag across all paragraph XML.
    # -----------------------------------------------------------------------
    try:
        all_xml = ''.join(p._element.xml for p in doc.paragraphs)
        n_inline = all_xml.count('wp:inline')
        n_anchor = all_xml.count('wp:anchor')
        if n_inline > 0:
            print(f"PASS: Component 5 — Image uses wp:inline ({n_inline} occurrence(s)); "
                  "None/inline wrapping confirmed (0.10 pts)")
            total_score += 0.10
        elif n_anchor > 0:
            print(f"FAIL: Component 5 — Image uses wp:anchor ({n_anchor} occurrence(s)); "
                  "expected wp:inline (None wrapping)")
        else:
            print("FAIL: Component 5 — No drawing element found for wrapping check")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -----------------------------------------------------------------------
    # Final score
    # -----------------------------------------------------------------------
    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
