"""
Initial Setup: Thesis outline with plain text paragraphs (no numbering)
Task ID: writer_list_028
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_list_028'
# Task context says file is at ~/Desktop/thesis_outline.docx
OUTPUT = f'{WORKDIR}/Desktop/thesis_outline.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Set document title via core properties
    doc.core_properties.title = "Thesis Outline"

    # Seven plain text paragraphs in Default Paragraph Style
    # No numbering - agent must apply Roman numeral list formatting
    sections = [
        "Abstract",
        "Literature Review",
        "Theoretical Framework",
        "Research Methodology",
        "Data Analysis and Results",
        "Discussion",
        "Conclusion and Future Work",
    ]

    for section_name in sections:
        para = doc.add_paragraph()
        para.style = doc.styles['Normal']
        run = para.add_run(section_name)
        # Plain text, no bold, no numbering

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
