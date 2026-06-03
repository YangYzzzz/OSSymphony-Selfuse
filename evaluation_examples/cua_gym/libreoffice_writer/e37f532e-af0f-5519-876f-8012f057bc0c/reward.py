"""
Reward Script: Apply uppercase Roman numeral numbering to thesis outline sections
Task ID: writer_list_028
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4 pts): All 7 paragraphs have list numbering (numPr) applied
  Component 2 (0.4 pts): The numbering format used is upperRoman (I, II, III...)
  Component 3 (0.2 pts): All paragraphs use list level 0 (default level 1 indentation)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_list_028'

FILE_PATH = f'{WORKDIR}/Desktop/thesis_outline.docx'

EXPECTED_PARAGRAPHS = [
    'Abstract',
    'Literature Review',
    'Theoretical Framework',
    'Research Methodology',
    'Data Analysis and Results',
    'Discussion',
    'Conclusion and Future Work',
]

NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def get_num_info(para):
    """Extract numId and ilvl from a paragraph's numPr element. Returns (numId, ilvl) or (None, None)."""
    pPr = para._element.find(NS + 'pPr')
    if pPr is None:
        return None, None
    numPr = pPr.find(NS + 'numPr')
    if numPr is None:
        return None, None
    ilvl_el = numPr.find(NS + 'ilvl')
    numId_el = numPr.find(NS + 'numId')
    ilvl = int(ilvl_el.get(NS + 'val')) if ilvl_el is not None else None
    numId = int(numId_el.get(NS + 'val')) if numId_el is not None else None
    return numId, ilvl


def get_abstract_num_fmt(doc, numId):
    """Look up the numFmt of abstractNum referenced by numId. Returns the numFmt string or None."""
    try:
        numbering_part = doc.part.numbering_part
        num_xml_el = numbering_part._element

        # Find the <w:num w:numId="numId"> element
        abstract_num_id_val = None
        for num_el in num_xml_el.findall(NS + 'num'):
            if num_el.get(NS + 'numId') == str(numId):
                abs_id_el = num_el.find(NS + 'abstractNumId')
                if abs_id_el is not None:
                    abstract_num_id_val = abs_id_el.get(NS + 'val')
                break

        if abstract_num_id_val is None:
            return None

        # Find the <w:abstractNum w:abstractNumId="abstract_num_id_val"> element
        for abs_num_el in num_xml_el.findall(NS + 'abstractNum'):
            if abs_num_el.get(NS + 'abstractNumId') == abstract_num_id_val:
                # Check level 0 numFmt
                for lvl_el in abs_num_el.findall(NS + 'lvl'):
                    if lvl_el.get(NS + 'ilvl') == '0':
                        num_fmt_el = lvl_el.find(NS + 'numFmt')
                        if num_fmt_el is not None:
                            return num_fmt_el.get(NS + 'val')
        return None
    except Exception:
        return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: check that the 7 expected paragraphs exist
    para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    all_expected_present = all(t in para_texts for t in EXPECTED_PARAGRAPHS)
    if not all_expected_present:
        missing = [t for t in EXPECTED_PARAGRAPHS if t not in para_texts]
        print(f"PRECONDITION FAIL: Missing paragraphs: {missing}")
        print("REWARD: 0.0")
        return 0.0

    content_paras = [p for p in doc.paragraphs if p.text.strip() in EXPECTED_PARAGRAPHS]

    # Component 1: All 7 target paragraphs have numPr (list numbering) applied (0.4 points)
    # This FAILS on initial (no numPr) and PASSES on golden (numPr present).
    try:
        numbered_count = 0
        first_num_id = None
        for p in content_paras:
            numId, ilvl = get_num_info(p)
            if numId is not None and numId != 0:
                numbered_count += 1
                if first_num_id is None:
                    first_num_id = numId

        if numbered_count == 7:
            print(f"PASS: Component 1 — All 7 paragraphs have list numbering applied (numPr present) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Expected 7 paragraphs with numbering, found {numbered_count}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")
        first_num_id = None

    # Component 2: The numbering format is upperRoman (I, II, III...) (0.4 points)
    # This FAILS on initial (no numbering) and PASSES on golden (upperRoman).
    try:
        if first_num_id is not None:
            num_fmt = get_abstract_num_fmt(doc, first_num_id)
            if num_fmt == 'upperRoman':
                print(f"PASS: Component 2 — Numbering format is upperRoman (I, II, III...) (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 2 — Expected numFmt='upperRoman', found numFmt='{num_fmt}'")
        else:
            print("FAIL: Component 2 — No numbering applied, cannot check format")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All paragraphs use list level 0 (ilvl=0, default level 1 indentation) (0.2 points)
    # This FAILS on initial (no numPr at all) and PASSES on golden (ilvl=0).
    try:
        level0_count = 0
        for p in content_paras:
            numId, ilvl = get_num_info(p)
            if numId is not None and numId != 0 and ilvl == 0:
                level0_count += 1

        if level0_count == 7:
            print(f"PASS: Component 3 — All 7 paragraphs at list level 0 (default level 1 indentation) (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — Expected 7 paragraphs at level 0, found {level0_count}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
