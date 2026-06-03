"""
Reward Script: Remote Work Policy Document
Task ID: writer_hr_058
Domain: libreoffice_writer
Scoring:
  C1 (0.15) - Header contains policy number HR-POL-2026-015
  C2 (0.15) - Main title uses Heading 1 style and mentions Remote Work
  C3 (0.20) - Page 1 headings: Policy Purpose, Eligibility Criteria, Application Process
  C4 (0.20) - Page 2 headings: Equipment & Technology, Work Schedule Expectations,
               Communication Requirements, Performance Evaluation
  C5 (0.15) - Page break exists between page 1 and page 2 content
  C6 (0.15) - Body content exists under headings (not just headings with no text)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_058'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather all paragraph info for analysis
    paragraphs = doc.paragraphs
    if len(paragraphs) == 0:
        print("FAIL: Document has no paragraphs at all")
        print("REWARD: 0.0")
        return 0.0

    # ---- Component 1: Header contains policy number (0.15 points) ----
    try:
        header_text = ""
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                header_text += " ".join(p.text for p in section.header.paragraphs)
        if "HR-POL-2026-015" in header_text:
            print(f"PASS: Component 1 — Header contains 'HR-POL-2026-015' (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Header text is '{header_text}', expected 'HR-POL-2026-015'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---- Component 2: Main title is Heading 1 with Remote Work content (0.15 points) ----
    try:
        heading1_found = False
        for para in paragraphs:
            style_name = para.style.name if para.style else ""
            if "Heading 1" in style_name:
                text_lower = para.text.lower()
                if "remote work" in text_lower:
                    heading1_found = True
                    print(f"PASS: Component 2 — Found Heading 1 '{para.text}' (0.15 pts)")
                    total_score += 0.15
                    break
                else:
                    print(f"FAIL: Component 2 — Heading 1 found but text '{para.text}' doesn't mention 'Remote Work'")
        if not heading1_found:
            # Check if any heading 1 exists at all
            has_h1 = any("Heading 1" in (p.style.name if p.style else "") for p in paragraphs)
            if not has_h1:
                print(f"FAIL: Component 2 — No Heading 1 style paragraph found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---- Component 3: Page 1 Heading 2 sections (0.20 points) ----
    # Expected: Policy Purpose, Eligibility Criteria, Application Process
    try:
        heading2_texts = []
        for para in paragraphs:
            style_name = para.style.name if para.style else ""
            if "Heading 2" in style_name:
                heading2_texts.append(para.text.strip())

        page1_expected = ["Policy Purpose", "Eligibility Criteria", "Application Process"]
        page1_found = 0
        for expected in page1_expected:
            exp_lower = expected.lower()
            # Check if any heading2 contains this topic (case-insensitive, partial match)
            for h2 in heading2_texts:
                if exp_lower in h2.lower() or all(w in h2.lower() for w in exp_lower.split()):
                    page1_found += 1
                    break

        if page1_found == 3:
            print(f"PASS: Component 3 — All 3 Page 1 headings found: {page1_expected} (0.20 pts)")
            total_score += 0.20
        elif page1_found > 0:
            partial = round(0.20 * page1_found / 3, 2)
            print(f"PARTIAL: Component 3 — {page1_found}/3 Page 1 headings found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — None of the Page 1 headings found. Heading 2s: {heading2_texts}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ---- Component 4: Page 2 Heading 2 sections (0.20 points) ----
    # Expected: Equipment & Technology, Work Schedule Expectations,
    #           Communication Requirements, Performance Evaluation
    try:
        page2_expected = [
            "Equipment",       # Equipment & Technology (partial match)
            "Work Schedule",   # Work Schedule Expectations
            "Communication",   # Communication Requirements
            "Performance Evaluation"
        ]
        page2_found = 0
        for expected in page2_expected:
            exp_lower = expected.lower()
            for h2 in heading2_texts:
                if exp_lower in h2.lower():
                    page2_found += 1
                    break

        if page2_found == 4:
            print(f"PASS: Component 4 — All 4 Page 2 headings found (0.20 pts)")
            total_score += 0.20
        elif page2_found > 0:
            partial = round(0.20 * page2_found / 4, 2)
            print(f"PARTIAL: Component 4 — {page2_found}/4 Page 2 headings found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — None of the Page 2 headings found. Heading 2s: {heading2_texts}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ---- Component 5: Page break exists (0.15 points) ----
    # Check for explicit page breaks in runs OR page_break_before on paragraphs
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        page_break_count = 0
        for para in paragraphs:
            for run in para.runs:
                for br in run.element.findall('.//w:br', ns):
                    btype = br.attrib.get(f'{{{ns["w"]}}}type', 'line')
                    if btype == 'page':
                        page_break_count += 1
            if para.paragraph_format.page_break_before:
                page_break_count += 1

        if page_break_count >= 1:
            print(f"PASS: Component 5 — Found {page_break_count} page break(s) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — No page breaks found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # ---- Component 6: Body content under headings (0.15 points) ----
    # Verify that there are Normal-style paragraphs with actual text content
    # (not just headings with no body text)
    try:
        normal_with_text = 0
        for para in paragraphs:
            style_name = para.style.name if para.style else ""
            # Count paragraphs that are not headings and have substantive text
            if "Heading" not in style_name and len(para.text.strip()) > 20:
                normal_with_text += 1

        # We expect at least 7 substantive body paragraphs (at least one per heading section)
        if normal_with_text >= 7:
            print(f"PASS: Component 6 — {normal_with_text} substantive body paragraphs found (0.15 pts)")
            total_score += 0.15
        elif normal_with_text >= 3:
            partial = round(0.15 * normal_with_text / 7, 2)
            print(f"PARTIAL: Component 6 — Only {normal_with_text}/7+ body paragraphs ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 6 — Only {normal_with_text} body paragraphs found, expected at least 7")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
