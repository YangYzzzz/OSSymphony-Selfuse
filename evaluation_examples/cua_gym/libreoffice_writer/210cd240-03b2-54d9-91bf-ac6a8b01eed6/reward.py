"""
Reward Script: Configure header with document title field centered in italic Liberation Sans 10pt
Task ID: writer_fs_090
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Header contains a TITLE field code
  Component 2 (0.25): Header field text shows 'Strategic Business Plan 2024'
  Component 3 (0.20): Header paragraph is center-aligned
  Component 4 (0.15): Header text is italic
  Component 5 (0.15): Header font is Liberation Sans 10pt
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_090'
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one section with a header
    if len(doc.sections) == 0:
        print("FAIL: No sections in document")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]
    header = section.header

    if not header.paragraphs:
        print("FAIL: No paragraphs in header")
        print("REWARD: 0.0")
        return 0.0

    hdr_para = header.paragraphs[0]
    hdr_elem = hdr_para._element

    # Component 1: Header contains a TITLE field code (0.25 points)
    # The TITLE field is inserted via Insert > Field > DocInformation > Title
    # It appears in XML as instrText containing " TITLE "
    try:
        instr_texts = hdr_elem.findall('.//w:instrText', NS)
        has_title_field = False
        for instr in instr_texts:
            if instr.text and 'TITLE' in instr.text.upper():
                has_title_field = True
                break

        if has_title_field:
            print(f"PASS: Component 1 — TITLE field code found in header (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — No TITLE field code found in header instrText elements")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header field text shows 'Strategic Business Plan 2024' (0.25 points)
    # The rendered text of the TITLE field should match the document title
    try:
        header_text = hdr_para.text.strip()
        expected_title = "Strategic Business Plan 2024"
        if header_text == expected_title:
            print(f"PASS: Component 2 — Header text is '{header_text}' (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — Expected '{expected_title}', found '{header_text}'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header paragraph is center-aligned (0.20 points)
    try:
        alignment = hdr_para.paragraph_format.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            print(f"PASS: Component 3 — Header paragraph is centered (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — Expected CENTER alignment, found {alignment}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Header text is italic (0.15 points)
    # Check that runs containing actual text (non-empty) are italic
    try:
        runs = hdr_para.runs
        # Filter to runs that have text or are part of the field
        text_runs = [r for r in runs if r.text.strip()]
        if not text_runs:
            # If no text runs, check all runs for italic property
            all_runs_with_rpr = runs
        else:
            all_runs_with_rpr = text_runs

        if len(all_runs_with_rpr) > 0:
            all_italic = all(r.font.italic for r in all_runs_with_rpr)
            if all_italic:
                print(f"PASS: Component 4 — Header text runs are italic (0.15 pts)")
                total_score += 0.15
            else:
                italic_status = [(r.text, r.font.italic) for r in all_runs_with_rpr]
                print(f"FAIL: Component 4 — Not all runs are italic: {italic_status}")
        else:
            print(f"FAIL: Component 4 — No runs found in header to check italic")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Header font is Liberation Sans 10pt (0.15 points)
    # Check font name and size on the text-bearing run(s)
    try:
        runs = hdr_para.runs
        text_runs = [r for r in runs if r.text.strip()]
        if not text_runs:
            text_runs = runs

        if len(text_runs) > 0:
            font_ok = True
            for r in text_runs:
                fname = r.font.name
                fsize = r.font.size
                fsize_pt = fsize.pt if fsize else None

                if fname != "Liberation Sans":
                    print(f"FAIL: Component 5 — Font name is '{fname}', expected 'Liberation Sans'")
                    font_ok = False
                    break
                if fsize_pt != 10.0:
                    print(f"FAIL: Component 5 — Font size is {fsize_pt}pt, expected 10.0pt")
                    font_ok = False
                    break

            if font_ok:
                print(f"PASS: Component 5 — Font is Liberation Sans 10pt (0.15 pts)")
                total_score += 0.15
        else:
            print(f"FAIL: Component 5 — No runs found in header to check font")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
