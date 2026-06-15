"""
Initial Setup: Configure header with document title field
Task ID: writer_fs_090
Domain: libreoffice_writer

Creates a Writer document with:
- Document properties Title = 'Strategic Business Plan 2024'
- Header enabled but empty
- Realistic business plan body content
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_090'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_core_property_title(doc, title_text):
    """Set the document core property Title via XML."""
    cp = doc.core_properties
    cp.title = title_text


def create_initial():
    doc = Document()

    # --- Set document properties Title ---
    set_core_property_title(doc, 'Strategic Business Plan 2024')

    # --- Enable header (but leave it empty) ---
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    # Clear any default content - ensure header is empty
    for para in header.paragraphs:
        para.clear()

    # --- Page setup ---
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Body content: Realistic business plan ---

    # Title heading
    title_para = doc.add_heading('Strategic Business Plan 2024', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by Meridian Consulting Group')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    run.italic = True

    doc.add_paragraph()  # spacer

    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This strategic business plan outlines the key objectives, market analysis, '
        'and financial projections for Meridian Technologies Inc. for the fiscal year 2024. '
        'Our primary focus areas include expanding into the Asia-Pacific market, launching '
        'two new product lines, and achieving a 25% increase in annual recurring revenue.'
    )
    doc.add_paragraph(
        'The company has demonstrated consistent growth over the past three years, with '
        'revenue increasing from $12.4M in 2021 to $18.7M in 2023. This plan details our '
        'strategy to reach the $23.4M target by December 2024.'
    )

    # Market Analysis
    doc.add_heading('2. Market Analysis', level=1)

    doc.add_heading('2.1 Industry Overview', level=2)
    doc.add_paragraph(
        'The global enterprise software market is projected to reach $376 billion by 2025, '
        'growing at a CAGR of 11.3%. Key drivers include digital transformation initiatives, '
        'cloud migration, and the increasing demand for AI-powered business intelligence tools.'
    )

    doc.add_heading('2.2 Target Market Segments', level=2)
    doc.add_paragraph(
        'Our target segments include mid-market enterprises (500-2,000 employees) in the '
        'financial services, healthcare, and manufacturing sectors. These verticals represent '
        'approximately $48 billion in addressable market opportunity.'
    )

    # Add a table for market data
    doc.add_heading('2.3 Competitive Landscape', level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['Competitor', 'Market Share', 'Annual Revenue', 'Key Strength']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Nextera Solutions', '18.2%', '$67.3M', 'Enterprise integration'],
        ['CloudPeak Systems', '14.7%', '$54.1M', 'Scalability'],
        ['DataBridge Corp', '11.5%', '$42.8M', 'Analytics platform'],
        ['Meridian Technologies', '6.3%', '$18.7M', 'Vertical expertise'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # Financial Projections
    doc.add_heading('3. Financial Projections', level=1)
    doc.add_paragraph(
        'Based on current pipeline analysis and market expansion plans, we project the '
        'following financial performance for 2024:'
    )

    fin_table = doc.add_table(rows=5, cols=3)
    fin_table.style = 'Table Grid'
    fin_headers = ['Metric', 'Q1-Q2 2024', 'Q3-Q4 2024']
    for i, h in enumerate(fin_headers):
        cell = fin_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    fin_data = [
        ['Revenue', '$10.8M', '$12.6M'],
        ['Gross Margin', '72.4%', '74.1%'],
        ['Operating Expenses', '$7.2M', '$7.8M'],
        ['Net Income', '$1.56M', '$2.13M'],
    ]
    for r, row_data in enumerate(fin_data, 1):
        for c, val in enumerate(row_data):
            fin_table.cell(r, c).text = val

    # Strategic Initiatives
    doc.add_heading('4. Strategic Initiatives', level=1)
    doc.add_paragraph(
        'The following initiatives will drive our growth trajectory:'
    )
    initiatives = [
        'Asia-Pacific Market Entry: Establish regional office in Singapore by Q2 2024 '
        'with initial team of 12 sales and support professionals.',
        'Product Line Expansion: Launch MeridianAI Analytics Suite and MeridianConnect '
        'Integration Platform, targeting $3.2M in first-year revenue.',
        'Customer Success Program: Implement proactive customer health monitoring to '
        'reduce churn from 8.3% to below 5% annually.',
        'Talent Acquisition: Recruit 45 new employees across engineering, sales, and '
        'customer success departments.',
    ]
    for init in initiatives:
        doc.add_paragraph(init, style='List Bullet')

    # Risk Assessment
    doc.add_heading('5. Risk Assessment', level=1)
    doc.add_paragraph(
        'Key risks identified include: macroeconomic uncertainty affecting enterprise IT '
        'budgets, increased competition from well-funded startups, and potential delays in '
        'product development timelines. Mitigation strategies are detailed in Appendix B.'
    )

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
