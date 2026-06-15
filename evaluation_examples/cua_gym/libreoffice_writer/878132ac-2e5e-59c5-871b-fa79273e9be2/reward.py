"""
Reward Script: Set table column widths in test_scores.docx
Task ID: writer_tbl_017
Domain: libreoffice_writer
Scoring:
  Component 1: First column ('Name') width == 5 cm (2835 twips)   — 0.6 pts
  Component 2: Second column ('Score') width == 3 cm (1701 twips) — 0.4 pts
  Cell contents: checked as a precondition gate (not a scoring component)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_017'
FILE_PATH = f'{WORKDIR}/Desktop/test_scores.docx'

# Exact twip values: 1 cm = 567 twips (OOXML dxa units)
COL1_TARGET_TWIPS = 2835   # 5 cm
COL2_TARGET_TWIPS = 1701   # 3 cm
# Allow ±1 twip tolerance for rounding
TOLERANCE = 1

# Expected cell contents used as precondition gate only
EXPECTED_CELLS = [
    (0, 0, 'Name'),
    (0, 1, 'Score'),
    (1, 0, 'Oliver'),
    (1, 1, '87'),
    (2, 0, 'Sophie'),
    (2, 1, '93'),
    (3, 0, 'Liam'),
    (3, 1, '78'),
]


def get_col_width_twips(table, col_index):
    """
    Get column width in twips from tblGrid definition.
    Returns int or None if unavailable.
    """
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        return None
    cols = tblGrid.findall(qn('w:gridCol'))
    if col_index >= len(cols):
        return None
    w = cols[col_index].get(qn('w:w'))
    if w is None:
        return None
    return int(w)


def get_cell_width_twips(cell):
    """
    Get an individual cell's declared width in twips (dxa units).
    Returns (int, str) = (twips, type) or (None, None).
    """
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None, None
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        return None, None
    w = tcW.get(qn('w:w'))
    wtype = tcW.get(qn('w:type'))
    return (int(w) if w is not None else None), wtype


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition gate: load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: exactly 1 table with 4 rows and 2 columns
    try:
        if len(doc.tables) != 1:
            print(f"CRITICAL: Expected 1 table, found {len(doc.tables)}")
            print("REWARD: 0.0")
            return 0.0
        table = doc.tables[0]
        if len(table.rows) != 4 or len(table.columns) != 2:
            print(f"CRITICAL: Expected 4x2 table, found {len(table.rows)}x{len(table.columns)}")
            print("REWARD: 0.0")
            return 0.0
    except Exception as e:
        print(f"CRITICAL: Table structure check failed: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: cell contents must be intact (not a scoring component)
    try:
        for (ri, ci, expected) in EXPECTED_CELLS:
            actual = table.rows[ri].cells[ci].text.strip()
            if actual != expected:
                print(f"PRECONDITION FAIL: Cell [{ri}][{ci}] has '{actual}', expected '{expected}' — content corrupted")
                print("REWARD: 0.0")
                return 0.0
        print("PRECONDITION: Cell contents intact (not scored)")
    except Exception as e:
        print(f"CRITICAL: Cell contents check failed: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: First column ('Name') width == 5 cm (2835 twips) (0.6 points)
    # Verifies both tblGrid grid col width AND all cell-level widths in column 0.
    # Fails on initial_env (4320 twips / 7.619 cm), passes on golden_env (2835 twips / 5.000 cm).
    try:
        grid_w0 = get_col_width_twips(table, 0)
        grid_ok = (grid_w0 is not None and abs(grid_w0 - COL1_TARGET_TWIPS) <= TOLERANCE)

        cell_col0_values = []
        for i, row in enumerate(table.rows):
            w, wtype = get_cell_width_twips(row.cells[0])
            cell_col0_values.append((i, w, wtype))

        # All cells in column 0 must match the target width
        cell_col0_mismatches = [
            (ri, w, wtype) for (ri, w, wtype) in cell_col0_values
            if w is None or abs(w - COL1_TARGET_TWIPS) > TOLERANCE
        ]
        cell_col0_ok = (len(cell_col0_mismatches) == 0)

        if grid_ok and cell_col0_ok:
            print(f"PASS: Component 1 — Col 0 (Name) width = {grid_w0} twips = {grid_w0/567:.3f} cm (target 5.000 cm) (0.6 pts)")
            total_score += 0.6
        else:
            if not grid_ok:
                gw_str = f"{grid_w0} twips ({grid_w0/567:.3f} cm)" if grid_w0 is not None else "None"
                print(f"FAIL: Component 1 — tblGrid col 0 = {gw_str}, expected {COL1_TARGET_TWIPS} twips (5.000 cm)")
            if not cell_col0_ok:
                for (ri, w, wtype) in cell_col0_mismatches:
                    w_str = f"{w} twips ({w/567:.3f} cm)" if w is not None else "None"
                    print(f"  row={ri} col=0 mismatch: {w_str}, type={wtype}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Second column ('Score') width == 3 cm (1701 twips) (0.4 points)
    # Verifies both tblGrid grid col width AND all cell-level widths in column 1.
    # Fails on initial_env (4320 twips / 7.619 cm), passes on golden_env (1701 twips / 3.000 cm).
    try:
        grid_w1 = get_col_width_twips(table, 1)
        grid_ok2 = (grid_w1 is not None and abs(grid_w1 - COL2_TARGET_TWIPS) <= TOLERANCE)

        cell_col1_values = []
        for i, row in enumerate(table.rows):
            w, wtype = get_cell_width_twips(row.cells[1])
            cell_col1_values.append((i, w, wtype))

        # All cells in column 1 must match the target width
        cell_col1_mismatches = [
            (ri, w, wtype) for (ri, w, wtype) in cell_col1_values
            if w is None or abs(w - COL2_TARGET_TWIPS) > TOLERANCE
        ]
        cell_col1_ok = (len(cell_col1_mismatches) == 0)

        if grid_ok2 and cell_col1_ok:
            print(f"PASS: Component 2 — Col 1 (Score) width = {grid_w1} twips = {grid_w1/567:.3f} cm (target 3.000 cm) (0.4 pts)")
            total_score += 0.4
        else:
            if not grid_ok2:
                gw_str = f"{grid_w1} twips ({grid_w1/567:.3f} cm)" if grid_w1 is not None else "None"
                print(f"FAIL: Component 2 — tblGrid col 1 = {gw_str}, expected {COL2_TARGET_TWIPS} twips (3.000 cm)")
            if not cell_col1_ok:
                for (ri, w, wtype) in cell_col1_mismatches:
                    w_str = f"{w} twips ({w/567:.3f} cm)" if w is not None else "None"
                    print(f"  row={ri} col=1 mismatch: {w_str}, type={wtype}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = round(min(total_score, 1.0), 6)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
