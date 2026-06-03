"""
Initial Setup: Clinical trial report with TOC listing 7 headings (8th heading missing from TOC)
Task ID: writer_struct_048
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
DESKTOP = '/home/user/Desktop'
TASK_ID = 'writer_struct_048'
OUTPUT = f'{DESKTOP}/clinical_trial.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc_entry(doc, text, page_num, level=1):
    """Add a TOC entry paragraph with dots and page number."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    if level == 1:
        run = para.add_run(text)
        run.font.size = Pt(11)
    else:
        para.paragraph_format.left_indent = Inches(0.25)
        run = para.add_run(text)
        run.font.size = Pt(10)
    # Add tab stop with dot leader
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.0), leader=1)  # WD_TAB_LEADER.DOTS
    run2 = para.add_run(f'\t{page_num}')
    run2.font.size = Pt(11) if level == 1 else Pt(10)
    return para


def add_section_heading(doc, text, level=1):
    """Add a heading with appropriate style."""
    para = doc.add_heading(text, level=level)
    return para


def add_body_text(doc, text):
    """Add a normal body paragraph."""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    return para


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ============================================================
    # TITLE PAGE
    # ============================================================
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    title_run = title_para.add_run('Phase III Randomized Controlled Trial')
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = subtitle_para.add_run(
        'Efficacy and Safety of NovaMab-7 in Treatment-Resistant\nRheumatoid Arthritis'
    )
    sub_run.font.size = Pt(14)

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_para.add_run('Clinical Study Report — Final Version\nMarch 2025')
    date_run.font.size = Pt(12)

    sponsor_para = doc.add_paragraph()
    sponsor_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sponsor_run = sponsor_para.add_run(
        'Sponsor: Meridian BioTherapeutics, Inc.\nProtocol Number: MBT-2024-RA-301'
    )
    sponsor_run.font.size = Pt(11)

    doc.add_page_break()

    # ============================================================
    # TABLE OF CONTENTS (7 entries — 8th heading NOT included)
    # ============================================================
    toc_heading = doc.add_paragraph()
    toc_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    toc_run = toc_heading.add_run('Table of Contents')
    toc_run.bold = True
    toc_run.font.size = Pt(14)
    toc_heading.paragraph_format.space_after = Pt(12)

    # 7 TOC entries (Appendix: Supplementary Data is intentionally omitted)
    toc_entries = [
        ('1.  Executive Summary', 3),
        ('2.  Background and Rationale', 3),
        ('3.  Study Objectives and Endpoints', 4),
        ('4.  Study Design and Methods', 5),
        ('5.  Patient Demographics and Baseline Characteristics', 6),
        ('6.  Efficacy Results', 7),
        ('7.  Safety and Adverse Events', 8),
    ]

    for entry_text, page_num in toc_entries:
        add_toc_entry(doc, entry_text, page_num, level=1)

    doc.add_page_break()

    # ============================================================
    # SECTION 1: Executive Summary
    # ============================================================
    add_section_heading(doc, '1.  Executive Summary', level=1)
    add_body_text(doc,
        'This report presents the findings of a Phase III randomized, double-blind, '
        'placebo-controlled trial evaluating NovaMab-7 for the treatment of '
        'moderate-to-severe treatment-resistant rheumatoid arthritis (RA). The study '
        'enrolled 487 patients across 42 investigative sites in 14 countries between '
        'January 2023 and November 2024.'
    )
    add_body_text(doc,
        'Primary endpoint results demonstrated a statistically significant reduction '
        'in Disease Activity Score (DAS28-CRP) at Week 24. The NovaMab-7 group achieved '
        'a mean DAS28-CRP reduction of 2.8 points compared to 0.9 points in the placebo '
        'group (p < 0.001). The ACR50 response rate was 62.3% vs 18.7% for placebo.'
    )
    add_body_text(doc,
        'The safety profile was consistent with the known class effects of IL-17A '
        'inhibitors. No unexpected safety signals were identified. The most common '
        'treatment-emergent adverse events included upper respiratory tract infection '
        '(14.2%), injection site reactions (8.6%), and headache (6.1%).'
    )
    doc.add_page_break()

    # ============================================================
    # SECTION 2: Background and Rationale
    # ============================================================
    add_section_heading(doc, '2.  Background and Rationale', level=1)
    add_body_text(doc,
        'Rheumatoid arthritis affects approximately 1% of the global adult population, '
        'with an estimated 1.5 million patients in the United States alone. Despite '
        'advances in biologic therapies, up to 40% of patients fail to achieve adequate '
        'disease control with available treatments, representing a significant unmet '
        'medical need.'
    )
    add_body_text(doc,
        'NovaMab-7 is a fully humanized monoclonal antibody targeting the interleukin-17A '
        '(IL-17A) pathway with enhanced receptor binding affinity. Preclinical studies '
        'demonstrated superior inhibition of inflammatory mediators compared to first-'
        'generation IL-17 inhibitors. Phase I studies in healthy volunteers confirmed '
        'acceptable pharmacokinetics and tolerability up to doses of 600 mg.'
    )
    add_body_text(doc,
        'Phase II results (Protocol MBT-2022-RA-201) in 186 patients with active RA '
        'demonstrated dose-dependent improvements in ACR response rates at 12 and 24 '
        'weeks, with 300 mg subcutaneous every 2 weeks selected as the Phase III dose '
        'based on the benefit-risk profile.'
    )

    # ============================================================
    # SECTION 3: Study Objectives and Endpoints
    # ============================================================
    add_section_heading(doc, '3.  Study Objectives and Endpoints', level=1)
    add_body_text(doc,
        'The primary objective was to demonstrate superiority of NovaMab-7 300 mg SC Q2W '
        'over placebo in reducing disease activity as measured by DAS28-CRP score at '
        'Week 24 in patients with moderate-to-severe RA who have had inadequate responses '
        'to at least two prior conventional DMARDs.'
    )
    add_body_text(doc,
        'Key secondary objectives included: (1) proportion of patients achieving ACR20, '
        'ACR50, and ACR70 response criteria at Weeks 12 and 24; (2) change from baseline '
        'in Health Assessment Questionnaire-Disability Index (HAQ-DI) at Week 24; '
        '(3) proportion achieving DAS28-CRP remission (score < 2.6) at Week 24; and '
        '(4) radiographic progression as measured by van der Heijde modified Sharp score '
        'at Week 52.'
    )
    doc.add_page_break()

    # ============================================================
    # SECTION 4: Study Design and Methods
    # ============================================================
    add_section_heading(doc, '4.  Study Design and Methods', level=1)
    add_body_text(doc,
        'This was a Phase III, multicenter, randomized, double-blind, placebo-controlled, '
        'parallel-group study. Eligible patients were randomized 2:1 to receive NovaMab-7 '
        '300 mg or matching placebo administered subcutaneously every 2 weeks for 52 weeks. '
        'Randomization was stratified by geographic region, prior biologic use, and baseline '
        'DAS28-CRP score (< 5.1 vs ≥ 5.1).'
    )
    add_body_text(doc,
        'The study comprised three periods: a screening period of up to 4 weeks, a '
        '52-week double-blind treatment period, and a 12-week safety follow-up period. '
        'Patients who completed the 52-week treatment period were eligible to enter a '
        'separate open-label extension study (Protocol MBT-2024-RA-302).'
    )
    add_body_text(doc,
        'All patients provided written informed consent prior to study participation. '
        'The study protocol and all amendments were approved by institutional review '
        'boards or independent ethics committees at each site. The study was conducted '
        'in accordance with ICH E6(R2) Good Clinical Practice guidelines and the '
        'Declaration of Helsinki.'
    )
    doc.add_page_break()

    # ============================================================
    # SECTION 5: Patient Demographics and Baseline Characteristics
    # ============================================================
    add_section_heading(doc, '5.  Patient Demographics and Baseline Characteristics', level=1)
    add_body_text(doc,
        'A total of 487 patients were randomized: 325 to NovaMab-7 and 162 to placebo. '
        'The safety analysis set comprised all patients who received at least one dose of '
        'study treatment (n=487). The full analysis set (n=481) included patients with '
        'baseline and at least one post-baseline DAS28-CRP assessment.'
    )
    add_body_text(doc,
        'Baseline demographic and disease characteristics were well balanced between '
        'treatment groups. The overall study population had a mean age of 52.4 years '
        '(SD: 12.8), was predominantly female (76.4%), and had a mean disease duration '
        'of 8.7 years. Mean baseline DAS28-CRP was 5.9 (SD: 0.8), consistent with '
        'moderately-to-highly active disease.'
    )
    add_body_text(doc,
        'Prior treatment history included methotrexate (91.2%), hydroxychloroquine '
        '(44.8%), sulfasalazine (38.3%), leflunomide (27.6%), and at least one prior '
        'biologic DMARD (58.9%). The most common prior biologic classes were TNF '
        'inhibitors (49.3%) and JAK inhibitors (21.8%).'
    )
    doc.add_page_break()

    # ============================================================
    # SECTION 6: Efficacy Results
    # ============================================================
    add_section_heading(doc, '6.  Efficacy Results', level=1)
    add_body_text(doc,
        'The primary endpoint was met with high statistical significance. At Week 24, '
        'mean change from baseline in DAS28-CRP was -2.81 (SE: 0.11) in the NovaMab-7 '
        'group versus -0.87 (SE: 0.16) in the placebo group, yielding a treatment '
        'difference of -1.94 (95% CI: -2.26, -1.62; p < 0.0001).'
    )
    add_body_text(doc,
        'ACR response rates at Week 24 were substantially higher with NovaMab-7: '
        'ACR20 (74.2% vs 31.5%), ACR50 (62.3% vs 18.7%), and ACR70 (41.8% vs 9.3%), '
        'all p < 0.0001. DAS28-CRP remission was achieved by 28.6% of NovaMab-7 patients '
        'vs 4.9% on placebo (p < 0.0001). HAQ-DI improvement from baseline was -0.62 '
        '(NovaMab-7) vs -0.21 (placebo), p < 0.0001.'
    )
    add_body_text(doc,
        'Radiographic data at Week 52 showed mean change in modified Sharp score of '
        '0.41 (NovaMab-7) vs 1.87 (placebo), demonstrating significant inhibition of '
        'structural damage progression (p = 0.003). Sixty-eight percent of NovaMab-7 '
        'patients showed no radiographic progression versus 49% for placebo.'
    )
    doc.add_page_break()

    # ============================================================
    # SECTION 7: Safety and Adverse Events
    # ============================================================
    add_section_heading(doc, '7.  Safety and Adverse Events', level=1)
    add_body_text(doc,
        'Treatment-emergent adverse events (TEAEs) were reported in 71.7% of NovaMab-7 '
        'patients and 68.5% of placebo patients. The incidence of serious adverse events '
        'was 8.3% vs 9.3% respectively, with no statistically significant difference '
        'between groups. Three fatal events occurred during the study (2 NovaMab-7, '
        '1 placebo), none assessed as related to study treatment.'
    )
    add_body_text(doc,
        'The most common TEAEs in the NovaMab-7 group included upper respiratory tract '
        'infection (14.2%), injection site reactions (8.6%), headache (6.1%), '
        'nasopharyngitis (5.8%), and diarrhea (4.6%). Injection site reactions were '
        'predominantly mild-to-moderate and did not lead to discontinuation. Candida '
        'infections were reported in 3.1% of NovaMab-7 patients vs 0.6% on placebo, '
        'consistent with the IL-17A mechanism.'
    )
    add_body_text(doc,
        'Immunogenicity analysis revealed treatment-emergent anti-drug antibodies (ADAs) '
        'in 4.6% of NovaMab-7 patients at any post-baseline timepoint. Neutralizing '
        'antibodies were detected in 1.8% of patients. ADA-positive patients did not '
        'show apparent impact on clinical efficacy or safety outcomes.'
    )

    # ============================================================
    # APPENDIX HEADING (Heading 1, but NOT in the TOC — this is what needs to be added)
    # ============================================================
    doc.add_page_break()
    add_section_heading(doc, 'Appendix: Supplementary Data', level=1)
    add_body_text(doc,
        'This appendix contains supplementary data tables, figures, and analyses '
        'referenced in the main report. All supplementary materials are provided '
        'to support transparency and reproducibility of the reported findings.'
    )

    # Supplementary Table 1
    supp_heading = doc.add_paragraph()
    supp_run = supp_heading.add_run('Table S1. Patient Disposition')
    supp_run.bold = True
    supp_run.font.size = Pt(11)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers_row = table.rows[0].cells
    headers_row[0].text = 'Disposition Category'
    headers_row[1].text = 'NovaMab-7 (n=325)'
    headers_row[2].text = 'Placebo (n=162)'
    for cell in headers_row:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    disposition_data = [
        ('Completed treatment (Week 52)', '279 (85.8%)', '133 (82.1%)'),
        ('Discontinued — adverse event', '18 (5.5%)', '10 (6.2%)'),
        ('Discontinued — lack of efficacy', '14 (4.3%)', '15 (9.3%)'),
        ('Discontinued — withdrawal of consent', '14 (4.3%)', '4 (2.5%)'),
    ]
    for i, (cat, nm7, placebo) in enumerate(disposition_data, 1):
        row = table.rows[i].cells
        row[0].text = cat
        row[1].text = nm7
        row[2].text = placebo

    doc.add_paragraph()

    # Supplementary Figure description
    add_body_text(doc,
        'Figure S1 (not shown): Kaplan-Meier curves for time to first ACR50 response '
        'in the full analysis set. Median time to ACR50 response was 8.3 weeks (95% '
        'CI: 7.1, 9.8) for NovaMab-7 versus not reached for placebo.'
    )

    add_body_text(doc,
        'Figure S2 (not shown): Forest plot of ACR50 response rates by subgroup '
        'at Week 24, including stratification by prior biologic use, geographic region, '
        'baseline DAS28-CRP, age group, sex, and disease duration.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
