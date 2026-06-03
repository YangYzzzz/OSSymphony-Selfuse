"""
Initial Setup: Create a letterhead document with an inline company logo image.
Task ID: writer_frd_070
Domain: libreoffice_writer
The logo is inserted as an inline image (wrap = None) in the document body.
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_070'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
LOGO_PATH = f'{WORKDIR}/company_logo.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_logo():
    """Create a semi-transparent company logo PNG."""
    from PIL import Image, ImageDraw, ImageFont

    width, height = 600, 400
    img = Image.new('RGBA', (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    # Draw a large circle as part of logo
    draw.ellipse([100, 50, 500, 350], fill=(30, 60, 120, 80), outline=(30, 60, 120, 120), width=3)

    # Draw company initials
    try:
        font_large = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 80)
        font_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 28)
    except OSError:
        font_large = ImageFont.load_default()
        font_small = ImageFont.load_default()

    draw.text((200, 120), "NV", fill=(30, 60, 120, 100), font=font_large)
    draw.text((160, 230), "NovaTech Inc.", fill=(30, 60, 120, 90), font=font_small)

    img.save(LOGO_PATH, 'PNG')
    print(f'Logo created: {LOGO_PATH}')


def create_initial():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    create_logo()

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Company header text ---
    header_para = doc.add_paragraph()
    header_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_para.paragraph_format.space_after = Pt(4)
    run = header_para.add_run("NovaTech Industries, Inc.")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(30, 60, 120)

    sub_para = doc.add_paragraph()
    sub_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_para.paragraph_format.space_after = Pt(2)
    run = sub_para.add_run("1200 Innovation Boulevard, Suite 400  |  San Francisco, CA 94107")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_para.paragraph_format.space_after = Pt(16)
    run = contact_para.add_run("Tel: (415) 555-0192  |  Fax: (415) 555-0193  |  www.novatechinc.com")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # --- Insert company logo image (inline / no wrap) ---
    logo_para = doc.add_paragraph()
    logo_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    logo_para.paragraph_format.space_after = Pt(20)
    run = logo_para.add_run()
    run.add_picture(LOGO_PATH, width=Inches(4.0))

    # --- Date line ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(12)
    run = date_para.add_run("March 28, 2026")
    run.font.size = Pt(11)

    # --- Recipient ---
    lines = [
        "Ms. Elena Vasquez",
        "Director of Operations",
        "Meridian Global Solutions",
        "4500 Pacific Coast Highway",
        "Long Beach, CA 90802",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(11)

    # Spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

    # --- Subject ---
    subj = doc.add_paragraph()
    subj.paragraph_format.space_after = Pt(12)
    run = subj.add_run("Re: Strategic Partnership Agreement — Q2 2026 Renewal")
    run.bold = True
    run.font.size = Pt(11)

    # --- Salutation ---
    sal = doc.add_paragraph()
    sal.paragraph_format.space_after = Pt(6)
    run = sal.add_run("Dear Ms. Vasquez,")
    run.font.size = Pt(11)

    # --- Body paragraphs ---
    body_texts = [
        "Thank you for your continued partnership with NovaTech Industries. We are pleased to "
        "present the terms for the renewal of our Strategic Partnership Agreement for the second "
        "quarter of 2026. As you know, our collaboration over the past eighteen months has yielded "
        "exceptional results, including a 23% increase in operational efficiency and a combined "
        "cost savings of $1.4 million across both organizations.",

        "Following the successful completion of Phase II of the Digital Transformation Initiative, "
        "our engineering team has identified several new areas where we can deepen our integration. "
        "Specifically, we propose expanding the scope of shared analytics infrastructure to include "
        "real-time supply chain monitoring and predictive demand forecasting. These capabilities "
        "would leverage the proprietary NovaTech Insight Platform, which was recently upgraded to "
        "support multi-tenant deployments.",

        "Enclosed with this letter you will find the updated Statement of Work (SOW), the revised "
        "pricing schedule reflecting the volume discount of 12% for the extended commitment period, "
        "and the technical integration roadmap prepared by our Solutions Architecture group. We have "
        "also included a summary of key performance indicators from the current agreement period "
        "for your review.",

        "We look forward to discussing these terms at your earliest convenience. Please do not "
        "hesitate to reach out to me directly at (415) 555-0192 or via email at "
        "j.martinez@novatechinc.com to schedule a meeting.",
    ]
    for text in body_texts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.font.size = Pt(11)

    # --- Closing ---
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    closing.paragraph_format.space_after = Pt(4)
    run = closing.add_run("Warm regards,")
    run.font.size = Pt(11)

    # Signature space
    sig_space = doc.add_paragraph()
    sig_space.paragraph_format.space_after = Pt(24)

    sig_name = doc.add_paragraph()
    sig_name.paragraph_format.space_after = Pt(0)
    run = sig_name.add_run("Javier Martinez")
    run.bold = True
    run.font.size = Pt(11)

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_after = Pt(0)
    run = sig_title.add_run("Vice President, Strategic Partnerships")
    run.font.size = Pt(11)

    sig_company = doc.add_paragraph()
    run = sig_company.add_run("NovaTech Industries, Inc.")
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
