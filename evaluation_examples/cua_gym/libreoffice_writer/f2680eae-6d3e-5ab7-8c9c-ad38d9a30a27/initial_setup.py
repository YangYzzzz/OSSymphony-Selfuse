"""
Initial Setup: Create a Writer document with heading structure for AutoText ToC task
Task ID: writer_fp_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_041'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Quarterly Business Report', level=0)

    # Heading 1: Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This quarterly report provides an overview of the company\'s performance '
        'across all divisions for Q1 2025. Revenue targets were exceeded by 12%, '
        'driven primarily by strong growth in the Asia-Pacific region and successful '
        'launch of three new product lines.'
    )
    doc.add_paragraph(
        'Key highlights include a 15% increase in customer acquisition, reduction '
        'in operational costs by 8%, and expansion into two new international markets. '
        'Employee satisfaction scores also improved by 6 points compared to last quarter.'
    )

    # Heading 1: Financial Performance
    doc.add_heading('Financial Performance', level=1)

    # Heading 2: Revenue Analysis
    doc.add_heading('Revenue Analysis', level=2)
    doc.add_paragraph(
        'Total revenue for Q1 2025 reached $42.7 million, representing a 12% increase '
        'over the same period last year. The breakdown by region shows North America '
        'contributing 45%, Europe at 30%, and Asia-Pacific at 25%.'
    )

    # Heading 3: North America
    doc.add_heading('North America', level=3)
    doc.add_paragraph(
        'The North American market generated $19.2 million in revenue. Enterprise '
        'software licensing accounted for 60% of regional revenue, while consulting '
        'services contributed the remaining 40%. Notable client wins include Meridian '
        'Healthcare Systems and Pacific Coast Financial Group.'
    )

    # Heading 3: Europe
    doc.add_heading('Europe', level=3)
    doc.add_paragraph(
        'European operations brought in $12.8 million. The UK and Germany remain our '
        'strongest markets, though France showed remarkable 22% growth following the '
        'appointment of a new regional director and revised go-to-market strategy.'
    )

    # Heading 3: Asia-Pacific
    doc.add_heading('Asia-Pacific', level=3)
    doc.add_paragraph(
        'Asia-Pacific revenue reached $10.7 million, a 28% increase year-over-year. '
        'Japan and Australia led growth, with emerging demand from Southeast Asian '
        'markets including Singapore and Vietnam.'
    )

    # Heading 2: Cost Management
    doc.add_heading('Cost Management', level=2)
    doc.add_paragraph(
        'Operating expenses were reduced by 8% through a combination of process '
        'automation, vendor renegotiation, and strategic workforce optimization. '
        'The cost-to-revenue ratio improved from 0.72 to 0.66.'
    )

    # Heading 1: Product Development
    doc.add_heading('Product Development', level=1)

    # Heading 2: New Product Launches
    doc.add_heading('New Product Launches', level=2)
    doc.add_paragraph(
        'Three new products were successfully launched during Q1: CloudSync Pro, '
        'DataGuard Enterprise, and Analytics Dashboard 3.0. Combined, these products '
        'generated $3.4 million in their first quarter, exceeding initial projections by 40%.'
    )

    # Heading 2: Research Initiatives
    doc.add_heading('Research Initiatives', level=2)
    doc.add_paragraph(
        'The R&D team invested 2,400 engineering hours in AI integration features '
        'planned for Q3 release. Early beta testing with select clients has produced '
        'favorable feedback, with an average satisfaction score of 4.6 out of 5.'
    )

    # Heading 3: AI Integration Roadmap
    doc.add_heading('AI Integration Roadmap', level=3)
    doc.add_paragraph(
        'The phased AI integration plan includes natural language query support in Q3, '
        'automated report generation in Q4, and predictive analytics capabilities '
        'scheduled for Q1 2026. Budget allocation for AI development totals $5.2 million.'
    )

    # Heading 1: Human Resources
    doc.add_heading('Human Resources', level=1)

    # Heading 2: Recruitment
    doc.add_heading('Recruitment', level=2)
    doc.add_paragraph(
        'A total of 47 new employees were hired across departments, with Engineering '
        'receiving 22 new team members, Sales adding 15, and Marketing welcoming 10. '
        'The average time-to-hire decreased from 34 to 28 days.'
    )

    # Heading 2: Employee Development
    doc.add_heading('Employee Development', level=2)
    doc.add_paragraph(
        'The Learning and Development team delivered 128 training sessions, reaching '
        '89% of the workforce. Leadership development programs were expanded to include '
        'mid-level managers, with 35 participants completing the first cohort.'
    )

    # Heading 1: Outlook
    doc.add_heading('Outlook', level=1)
    doc.add_paragraph(
        'Based on current momentum and market conditions, Q2 2025 revenue is projected '
        'at $46.5 million. Key priorities include accelerating AI product development, '
        'expanding the partner ecosystem, and entering the Latin American market. '
        'The Board of Directors has approved a $3 million investment in market expansion '
        'activities for the remainder of the fiscal year.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
