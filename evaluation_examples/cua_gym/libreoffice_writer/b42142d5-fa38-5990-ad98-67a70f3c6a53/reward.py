"""
Reward Script: Job Posting for Human Resources Manager
Task ID: writer_hr_054
Domain: libreoffice_writer
Scoring:
  C1: Heading 1 title present (0.15)
  C2: Key Responsibilities section with 8 bullet points (0.20)
  C3: Required Qualifications section with 6 bullet points (0.20)
  C4: Preferred Qualifications section with 4 bullet points (0.15)
  C5: Compensation & Benefits section present (0.10)
  C6: How to Apply section present (0.10)
  C7: EOE statement in italic at bottom (0.10)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_054'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have content (initial is blank)
    if len(doc.paragraphs) == 0:
        print("FAIL: Document has no paragraphs (blank document)")
        print("REWARD: 0.0")
        return 0.0

    # Build a structured view of the document: headings and their bullet children
    heading_sections = {}  # heading_text -> list of bullet texts
    all_headings_h1 = []
    all_headings_h2 = []
    current_heading = None

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        if style_name == 'Heading 1' and para.text.strip():
            all_headings_h1.append(para.text.strip())
            current_heading = para.text.strip()
            heading_sections[current_heading] = []
        elif style_name == 'Heading 2' and para.text.strip():
            all_headings_h2.append(para.text.strip())
            current_heading = para.text.strip()
            heading_sections[current_heading] = []
        elif style_name in ('List Bullet', 'List Number') and current_heading:
            heading_sections[current_heading].append(para.text.strip())

    print(f"INFO: Found {len(all_headings_h1)} H1 headings, {len(all_headings_h2)} H2 headings")
    print(f"INFO: Sections: {list(heading_sections.keys())}")

    # Component 1: Heading 1 title present (0.15 points)
    # The job title should be in Heading 1 style
    try:
        if len(all_headings_h1) > 0:
            h1_text = all_headings_h1[0].lower()
            # Check that it references a job title related to HR
            if 'human resources' in h1_text or 'hr manager' in h1_text or 'hr ' in h1_text:
                print(f"PASS: Component 1 — Heading 1 title found: '{all_headings_h1[0]}' (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 1 — Heading 1 found but doesn't match HR Manager: '{all_headings_h1[0]}'")
        else:
            print("FAIL: Component 1 — No Heading 1 found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Helper to find a section by keyword matching
    def find_section(keywords, headings_dict):
        """Find a section whose heading contains any of the keywords (case-insensitive)."""
        for heading, bullets in headings_dict.items():
            h_lower = heading.lower()
            if any(kw in h_lower for kw in keywords):
                return heading, bullets
        return None, []

    # Component 2: Key Responsibilities section with 8 bullet points (0.20 points)
    try:
        heading, bullets = find_section(['key responsibilities', 'responsibilities'], heading_sections)
        if heading:
            count = len(bullets)
            if count >= 8:
                print(f"PASS: Component 2 — '{heading}' has {count} bullets (>= 8) (0.20 pts)")
                total_score += 0.20
            elif count >= 5:
                partial = 0.10
                print(f"PARTIAL: Component 2 — '{heading}' has {count}/8 bullets ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — '{heading}' has only {count}/8 bullets")
        else:
            print("FAIL: Component 2 — No 'Key Responsibilities' section found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Required Qualifications section with 6 bullet points (0.20 points)
    try:
        heading, bullets = find_section(['required qualifications'], heading_sections)
        if heading:
            count = len(bullets)
            if count >= 6:
                print(f"PASS: Component 3 — '{heading}' has {count} bullets (>= 6) (0.20 pts)")
                total_score += 0.20
            elif count >= 4:
                partial = 0.10
                print(f"PARTIAL: Component 3 — '{heading}' has {count}/6 bullets ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — '{heading}' has only {count}/6 bullets")
        else:
            print("FAIL: Component 3 — No 'Required Qualifications' section found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Preferred Qualifications section with 4 bullet points (0.15 points)
    try:
        heading, bullets = find_section(['preferred qualifications'], heading_sections)
        if heading:
            count = len(bullets)
            if count >= 4:
                print(f"PASS: Component 4 — '{heading}' has {count} bullets (>= 4) (0.15 pts)")
                total_score += 0.15
            elif count >= 2:
                partial = 0.07
                print(f"PARTIAL: Component 4 — '{heading}' has {count}/4 bullets ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — '{heading}' has only {count}/4 bullets")
        else:
            print("FAIL: Component 4 — No 'Preferred Qualifications' section found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Compensation & Benefits section present (0.10 points)
    try:
        heading, bullets = find_section(['compensation', 'benefits'], heading_sections)
        if heading:
            print(f"PASS: Component 5 — '{heading}' section found (0.10 pts)")
            total_score += 0.10
        else:
            print("FAIL: Component 5 — No 'Compensation & Benefits' section found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: How to Apply section present (0.10 points)
    try:
        heading, _ = find_section(['how to apply', 'application'], heading_sections)
        if heading:
            print(f"PASS: Component 6 — '{heading}' section found (0.10 pts)")
            total_score += 0.10
        else:
            print("FAIL: Component 6 — No 'How to Apply' section found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: EOE statement in italic at the bottom (0.10 points)
    # The last paragraph should contain "equal opportunity" text and be in italic
    try:
        # Search last two paragraphs for EOE statement
        eoe_para = None
        for idx in range(-1, max(-3, -len(doc.paragraphs) - 1), -1):
            candidate = doc.paragraphs[idx]
            ctext = candidate.text.lower()
            if 'equal opportunity' in ctext or 'eoe' in ctext or 'equal employment' in ctext:
                eoe_para = candidate
                break

        if eoe_para is not None:
            italic_runs = [r for r in eoe_para.runs if r.italic]
            if italic_runs:
                print(f"PASS: Component 7 — EOE statement found in italic at bottom (0.10 pts)")
                total_score += 0.10
            elif len(eoe_para.runs) > 0:
                print(f"PARTIAL: Component 7 — EOE statement found but not italic (0.05 pts)")
                total_score += 0.05
            else:
                print(f"FAIL: Component 7 — EOE paragraph has no runs to check italic")
        else:
            print(f"FAIL: Component 7 — No EOE statement found at bottom of document")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'

persist_app_state()

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
