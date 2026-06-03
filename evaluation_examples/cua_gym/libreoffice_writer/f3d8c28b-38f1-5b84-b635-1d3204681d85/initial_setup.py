"""
Initial Setup: Insert a text frame with a pull quote from a magazine article
Task ID: writer_rd_023
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_023'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading("The Future of Technology: Trends Shaping Tomorrow", level=1)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(18)

    # Subtitle / byline
    byline = doc.add_paragraph()
    byline.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    byline.paragraph_format.space_after = Pt(24)
    run = byline.add_run("By Priya Nakamura  |  Technology Correspondent  |  March 2025")
    run.font.size = Pt(11)
    run.italic = True

    # Paragraph 1
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(10)
    run = p1.add_run(
        "The technology landscape is evolving at an unprecedented pace, reshaping industries "
        "and redefining how we interact with the world around us. From artificial intelligence "
        "breakthroughs to quantum computing milestones, the innovations of the past decade "
        "have laid the groundwork for a transformative era that promises to touch every facet "
        "of human life."
    )
    run.font.size = Pt(11)

    # Paragraph 2
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    run = p2.add_run(
        "Artificial intelligence continues to dominate headlines, but beyond the buzz, "
        "real-world applications are quietly revolutionizing healthcare, logistics, and "
        "education. Diagnostic algorithms now detect early-stage cancers with accuracy "
        "rates surpassing those of experienced radiologists, while predictive models "
        "optimize supply chains across global networks serving billions of consumers."
    )
    run.font.size = Pt(11)

    # Paragraph 3
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(10)
    run = p3.add_run(
        "Innovation is not just about technology; it is about solving real problems. "
        "This philosophy has driven startups and established firms alike to focus on "
        "outcomes rather than features. Companies that align their engineering efforts "
        "with genuine user needs consistently outperform those that chase the latest "
        "technical trends without a clear purpose."
    )
    run.font.size = Pt(11)

    # Paragraph 4
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(10)
    run = p4.add_run(
        "The rise of edge computing represents another seismic shift. By processing data "
        "closer to its source, edge architectures dramatically reduce latency and bandwidth "
        "costs. Autonomous vehicles, smart factories, and augmented reality applications all "
        "depend on the sub-millisecond response times that only edge infrastructure can "
        "reliably deliver at scale."
    )
    run.font.size = Pt(11)

    # Paragraph 5
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(10)
    run = p5.add_run(
        "Cybersecurity has emerged as one of the most critical challenges of the digital age. "
        "With the average cost of a data breach now exceeding $4.5 million, organizations are "
        "investing heavily in zero-trust architectures, advanced threat detection powered by "
        "machine learning, and employee training programs designed to mitigate the human "
        "factor in security vulnerabilities."
    )
    run.font.size = Pt(11)

    # Paragraph 6
    p6 = doc.add_paragraph()
    p6.paragraph_format.space_after = Pt(10)
    run = p6.add_run(
        "Sustainability has become a core pillar of technology strategy. Major cloud providers "
        "are committing to carbon-neutral operations by 2030, while semiconductor manufacturers "
        "invest in energy-efficient chip designs that deliver more computational power per watt "
        "than ever before. Green computing is no longer a niche concern; it is a competitive "
        "imperative driving investment decisions across the industry."
    )
    run.font.size = Pt(11)

    # Paragraph 7
    p7 = doc.add_paragraph()
    p7.paragraph_format.space_after = Pt(10)
    run = p7.add_run(
        "Quantum computing, once confined to theoretical physics departments, is inching "
        "toward practical applications. Recent demonstrations of quantum advantage in "
        "molecular simulation and optimization problems suggest that the technology could "
        "revolutionize drug discovery, materials science, and financial modeling within "
        "the next five to ten years."
    )
    run.font.size = Pt(11)

    # Paragraph 8
    p8 = doc.add_paragraph()
    p8.paragraph_format.space_after = Pt(10)
    run = p8.add_run(
        "As these technologies converge, the organizations that will thrive are those "
        "that invest not only in tools but in talent and culture. Building diverse, "
        "interdisciplinary teams and fostering a mindset of continuous learning will be "
        "essential for navigating the complexity of tomorrow's technological landscape. "
        "The future belongs to those who prepare for it today."
    )
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
