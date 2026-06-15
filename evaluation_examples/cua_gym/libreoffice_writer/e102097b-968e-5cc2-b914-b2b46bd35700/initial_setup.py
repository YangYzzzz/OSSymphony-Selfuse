"""
Initial Setup: Writer document with version history
Task ID: writer_lec_073
Domain: libreoffice_writer

Creates .odt with 4 versions. Version 2 labeled 'Before reformatting'.
Current doc has formatting changes applied after version 2.
Uses docx->odt conversion for proper ODF base, then ZIP injection for versions.
"""
import os
import shlex
import subprocess
import time
import zipfile
import io
import xml.etree.ElementTree as ET

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_073'
OUTPUT = f'{WORKDIR}/{TASK_ID}.odt'

def launch_gui(command, delay_sec=1.0):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(shlex.split(command),
                     stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, env=env)
    time.sleep(delay_sec)

def kill_lo():
    subprocess.run(["pkill", "-9", "-f", "soffice"], capture_output=True)
    time.sleep(3)

def create_docx_and_convert(docx_path, odt_path):
    """Create a docx then convert to odt using LO headless."""
    result = subprocess.run(
        ["libreoffice", "--headless", "--norestore", "--convert-to", "odt",
         "--outdir", os.path.dirname(odt_path) or "/tmp", docx_path],
        capture_output=True, text=True, timeout=30)
    if "writer8" not in result.stdout and "writer_pdf_Export" not in result.stdout:
        # Check if file exists
        expected = os.path.join(os.path.dirname(odt_path) or "/tmp",
                               os.path.basename(docx_path).replace('.docx', '.odt'))
        if os.path.exists(expected):
            if expected != odt_path:
                os.rename(expected, odt_path)
        else:
            print(f"WARNING: Conversion may have failed: {result.stdout} {result.stderr}")

def create_initial():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    kill_lo()

    ########################################
    # Create Version 1 content (Initial draft - just title + intro)
    ########################################
    v1_doc = Document()
    v1_doc.add_heading("Annual Report 2025", level=1)
    v1_doc.add_paragraph(
        "This document presents the annual performance review for fiscal year 2025. "
        "Our organization achieved significant milestones across all departments.")
    v1_docx = "/tmp/v1.docx"
    v1_odt = "/tmp/v1.odt"
    v1_doc.save(v1_docx)
    create_docx_and_convert(v1_docx, v1_odt)

    ########################################
    # Create Version 2 content (Before reformatting - full content, plain format)
    ########################################
    v2_doc = Document()
    v2_doc.add_heading("Annual Report 2025", level=1)

    v2_doc.add_heading("Executive Summary", level=2)
    v2_doc.add_paragraph(
        "This document presents the annual performance review for fiscal year 2025. "
        "Our organization achieved significant milestones across all departments, "
        "exceeding revenue targets and expanding into three new markets.")

    v2_doc.add_heading("Financial Overview", level=2)
    v2_doc.add_paragraph(
        "Total revenue increased by 15% compared to the previous fiscal year, "
        "reaching $47.2 million. Operating expenses remained within budget at "
        "$2.3 million, resulting in a net profit margin of 22%.")

    v2_doc.add_heading("Department Highlights", level=2)
    v2_doc.add_paragraph(
        "Engineering delivered 12 major product releases on schedule. "
        "The platform uptime exceeded 99.97% throughout the year.")
    v2_doc.add_paragraph(
        "Marketing campaigns generated 45,000 new leads, a 30% increase over 2024. "
        "Brand awareness surveys showed a 12-point improvement.")
    v2_doc.add_paragraph(
        "Customer support maintained a 98.5% satisfaction rating while handling "
        "23% more tickets than the previous year.")

    v2_doc.add_heading("Outlook for 2026", level=2)
    v2_doc.add_paragraph(
        "We project continued growth with a focus on international expansion "
        "and product diversification. Key initiatives include opening offices "
        "in Berlin and Tokyo, and launching our enterprise tier.")

    v2_docx = "/tmp/v2.docx"
    v2_odt = "/tmp/v2.odt"
    v2_doc.save(v2_docx)
    create_docx_and_convert(v2_docx, v2_odt)

    ########################################
    # Create Version 3 content (V2 + appendix)
    ########################################
    v3_doc = Document()
    v3_doc.add_heading("Annual Report 2025", level=1)
    v3_doc.add_heading("Executive Summary", level=2)
    v3_doc.add_paragraph(
        "This document presents the annual performance review for fiscal year 2025. "
        "Our organization achieved significant milestones across all departments, "
        "exceeding revenue targets and expanding into three new markets.")
    v3_doc.add_heading("Financial Overview", level=2)
    v3_doc.add_paragraph(
        "Total revenue increased by 15% compared to the previous fiscal year, "
        "reaching $47.2 million. Operating expenses remained within budget at "
        "$2.3 million, resulting in a net profit margin of 22%.")
    v3_doc.add_heading("Department Highlights", level=2)
    v3_doc.add_paragraph(
        "Engineering delivered 12 major product releases on schedule. "
        "The platform uptime exceeded 99.97% throughout the year.")
    v3_doc.add_paragraph(
        "Marketing campaigns generated 45,000 new leads, a 30% increase over 2024. "
        "Brand awareness surveys showed a 12-point improvement.")
    v3_doc.add_paragraph(
        "Customer support maintained a 98.5% satisfaction rating while handling "
        "23% more tickets than the previous year.")
    v3_doc.add_heading("Outlook for 2026", level=2)
    v3_doc.add_paragraph(
        "We project continued growth with a focus on international expansion "
        "and product diversification. Key initiatives include opening offices "
        "in Berlin and Tokyo, and launching our enterprise tier.")
    v3_doc.add_heading("Appendix A: Quarterly Revenue Breakdown", level=2)
    v3_doc.add_paragraph("Q1: $10.8M | Q2: $11.2M | Q3: $12.1M | Q4: $13.1M")
    v3_doc.add_paragraph(
        "The strongest performance was in Q4, driven by holiday season sales "
        "and new enterprise contracts signed in November.")

    v3_docx = "/tmp/v3.docx"
    v3_odt = "/tmp/v3.odt"
    v3_doc.save(v3_docx)
    create_docx_and_convert(v3_docx, v3_odt)

    ########################################
    # Create CURRENT state (V3 content + formatting changes)
    ########################################
    current_doc = Document()

    # Title - centered, colored, large
    title = current_doc.add_heading("Annual Report 2025", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        run.font.size = Pt(24)

    def add_formatted_heading(doc, text):
        h = doc.add_heading(text, level=2)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            run.font.size = Pt(18)
        return h

    def add_formatted_body(doc, text):
        p = doc.add_paragraph(text)
        pf = p.paragraph_format
        pf.left_indent = Inches(0.3)
        pf.line_spacing = 1.5
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Arial"
        return p

    add_formatted_heading(current_doc, "Executive Summary")
    p = add_formatted_body(current_doc,
        "This document presents the annual performance review for fiscal year 2025. "
        "Our organization achieved significant milestones across all departments, "
        "exceeding revenue targets and expanding into three new markets.")

    add_formatted_heading(current_doc, "Financial Overview")
    p = add_formatted_body(current_doc,
        "Total revenue increased by 15% compared to the previous fiscal year, "
        "reaching $47.2 million. Operating expenses remained within budget at "
        "$2.3 million, resulting in a net profit margin of 22%.")

    add_formatted_heading(current_doc, "Department Highlights")
    add_formatted_body(current_doc,
        "Engineering delivered 12 major product releases on schedule. "
        "The platform uptime exceeded 99.97% throughout the year.")
    add_formatted_body(current_doc,
        "Marketing campaigns generated 45,000 new leads, a 30% increase over 2024. "
        "Brand awareness surveys showed a 12-point improvement.")
    add_formatted_body(current_doc,
        "Customer support maintained a 98.5% satisfaction rating while handling "
        "23% more tickets than the previous year.")

    add_formatted_heading(current_doc, "Outlook for 2026")
    add_formatted_body(current_doc,
        "We project continued growth with a focus on international expansion "
        "and product diversification. Key initiatives include opening offices "
        "in Berlin and Tokyo, and launching our enterprise tier.")

    add_formatted_heading(current_doc, "Appendix A: Quarterly Revenue Breakdown")
    add_formatted_body(current_doc,
        "Q1: $10.8M | Q2: $11.2M | Q3: $12.1M | Q4: $13.1M")
    add_formatted_body(current_doc,
        "The strongest performance was in Q4, driven by holiday season sales "
        "and new enterprise contracts signed in November.")

    current_docx = "/tmp/current.docx"
    current_odt = "/tmp/current.odt"
    current_doc.save(current_docx)
    create_docx_and_convert(current_docx, current_odt)

    print("All ODT files created via conversion")

    ########################################
    # Now build the final versioned ODF
    ########################################

    # Read version files as binary
    with open(v1_odt, 'rb') as f:
        v1_data = f.read()
    with open(v2_odt, 'rb') as f:
        v2_data = f.read()
    with open(v3_odt, 'rb') as f:
        v3_data = f.read()
    v4_data = v3_data  # V4 same as V3 (pre-formatting snapshot)

    # Read current (formatted) document as base
    with open(current_odt, 'rb') as f:
        base_data = f.read()

    base_zip = zipfile.ZipFile(io.BytesIO(base_data), 'r')

    # VersionList.xml
    version_list_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<VL:version-list xmlns:VL="http://openoffice.org/2001/version-list"
    xmlns:dc="http://purl.org/dc/elements/1.1/">
  <VL:version-entry VL:title="Version1"
      VL:comment="Initial draft"
      VL:creator="Sarah Mitchell"
      dc:date-time="2025-10-15T09:30:00"/>
  <VL:version-entry VL:title="Version2"
      VL:comment="Before reformatting"
      VL:creator="Sarah Mitchell"
      dc:date-time="2025-10-22T14:15:00"/>
  <VL:version-entry VL:title="Version3"
      VL:comment="Added appendix with quarterly data"
      VL:creator="Sarah Mitchell"
      dc:date-time="2025-11-05T10:45:00"/>
  <VL:version-entry VL:title="Version4"
      VL:comment="Pre-formatting snapshot"
      VL:creator="Sarah Mitchell"
      dc:date-time="2025-11-12T16:00:00"/>
</VL:version-list>'''

    # Update manifest to include version entries
    ns_m = 'urn:oasis:names:tc:opendocument:xmlns:manifest:1.0'
    ET.register_namespace('manifest', ns_m)
    manifest = ET.fromstring(base_zip.read('META-INF/manifest.xml').decode())

    # Add VersionList.xml entry
    e = ET.SubElement(manifest, f'{{{ns_m}}}file-entry')
    e.set(f'{{{ns_m}}}media-type', 'text/xml')
    e.set(f'{{{ns_m}}}full-path', 'VersionList.xml')

    # Add version data entries
    for i in range(1, 5):
        e = ET.SubElement(manifest, f'{{{ns_m}}}file-entry')
        e.set(f'{{{ns_m}}}media-type', 'application/vnd.oasis.opendocument.text')
        e.set(f'{{{ns_m}}}full-path', f'Versions/Version{i}')

    new_manifest = ET.tostring(manifest, encoding='unicode', xml_declaration=True)

    # Build final ODF
    with open(OUTPUT, 'wb') as f:
        with zipfile.ZipFile(f, 'w', zipfile.ZIP_DEFLATED) as zf:
            for item in base_zip.namelist():
                if item == 'META-INF/manifest.xml':
                    zf.writestr(item, new_manifest)
                else:
                    data = base_zip.read(item)
                    if item == 'mimetype':
                        zf.writestr(item, data, compress_type=zipfile.ZIP_STORED)
                    else:
                        zf.writestr(item, data)

            # Add version list and data
            zf.writestr('VersionList.xml', version_list_xml)
            zf.writestr('Versions/Version1', v1_data)
            zf.writestr('Versions/Version2', v2_data)
            zf.writestr('Versions/Version3', v3_data)
            zf.writestr('Versions/Version4', v4_data)

    base_zip.close()

    # Verify the file loads in LO
    result = subprocess.run(
        ["libreoffice", "--headless", "--norestore", "--convert-to", "pdf",
         "--outdir", "/tmp", OUTPUT],
        capture_output=True, text=True, timeout=30)
    if "writer_pdf_Export" in result.stdout:
        print(f"VERIFIED: {OUTPUT} loads successfully in LibreOffice")
    else:
        print(f"WARNING: File may not load correctly: {result.stdout} {result.stderr}")

    # Verify ZIP structure
    with zipfile.ZipFile(OUTPUT, 'r') as z:
        names = z.namelist()
        ver_entries = [n for n in names if n.startswith("Versions/")]
        print(f"ZIP entries: {len(names)}")
        print(f"Version entries: {ver_entries}")
        print(f"VersionList.xml present: {'VersionList.xml' in names}")
        print(f"File size: {os.path.getsize(OUTPUT)} bytes")

    # Clean up temp files
    for p in [v1_docx, v1_odt, v2_docx, v2_odt, v3_docx, v3_odt,
              current_docx, current_odt, "/tmp/v1.odt", "/tmp/v2.odt",
              "/tmp/v3.odt", "/tmp/current.odt"]:
        try:
            os.remove(p)
        except:
            pass

    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    # Use pyautogui to click Yes on repair dialog if it appears
    kill_lo()
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=5.0)

    # Handle potential repair dialog
    try:
        import pyautogui
        time.sleep(3)
        # Check for repair dialog and click Yes
        screenshot = pyautogui.screenshot()
        # Look for "Yes" button - it should be near center-right of screen
        # The repair dialog is centered, Yes button is on the right
        try:
            yes_loc = pyautogui.locateOnScreen
        except:
            pass
        # Just try clicking where Yes would be (approximate)
        # Based on the screenshot: Yes button was at approximately (480/800 * 1920, 285/600 * 1080)
        # = (1152, 513)
        pyautogui.click(1152, 513)
        time.sleep(2)
    except Exception as e:
        print(f"PyAutoGUI handling: {e}")

    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
