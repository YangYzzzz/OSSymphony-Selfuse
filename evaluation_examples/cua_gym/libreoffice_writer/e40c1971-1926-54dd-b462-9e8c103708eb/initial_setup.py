"""
Initial Setup: Insert a date/time field in the footer
Task ID: writer_tm_068
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_068'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Footer must be EMPTY (task is to add date/time field to footer)

    # --- Document Title ---
    title = doc.add_heading('Meeting Minutes - Q1 Strategy Review', level=1)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Meeting Details ---
    details = doc.add_paragraph()
    details.paragraph_format.space_after = Pt(6)
    run = details.add_run('Date: ')
    run.bold = True
    details.add_run('March 15, 2026')
    details.add_run('    ')
    run2 = details.add_run('Time: ')
    run2.bold = True
    details.add_run('10:00 AM - 11:30 AM')
    details.add_run('    ')
    run3 = details.add_run('Location: ')
    run3.bold = True
    details.add_run('Conference Room B, 4th Floor')

    attendees = doc.add_paragraph()
    attendees.paragraph_format.space_after = Pt(12)
    run_a = attendees.add_run('Attendees: ')
    run_a.bold = True
    attendees.add_run('Sarah Chen (VP Engineering), Marcus Johnson (Product Lead), '
                      'Elena Rodriguez (Design Director), David Kim (QA Manager), '
                      'Priya Patel (Marketing Head), James O\'Brien (CFO)')

    # --- Section 1: Opening Remarks ---
    doc.add_heading('1. Opening Remarks', level=2)
    doc.add_paragraph(
        'Sarah Chen opened the meeting at 10:02 AM, welcoming all participants and '
        'noting that this quarterly review covers the period from January through March 2026. '
        'She highlighted three key objectives for the session: reviewing product milestones, '
        'assessing budget utilization, and aligning on Q2 priorities.'
    )

    # --- Section 2: Product Update ---
    doc.add_heading('2. Product Development Update', level=2)
    doc.add_paragraph(
        'Marcus Johnson presented the Q1 product roadmap progress. The team successfully '
        'launched version 3.2 of the platform on February 20th, which included the new '
        'dashboard analytics module, improved search functionality, and the customer feedback '
        'integration pipeline. User adoption metrics show a 23% increase in daily active users '
        'compared to Q4 2025.'
    )
    doc.add_paragraph(
        'Key milestones achieved during Q1 include:'
    )
    doc.add_paragraph('Completion of the API v2 migration affecting 847 endpoints', style='List Bullet')
    doc.add_paragraph('Reduction in average page load time from 2.3s to 1.1s', style='List Bullet')
    doc.add_paragraph('Deployment of automated testing framework covering 92% of codebase', style='List Bullet')
    doc.add_paragraph('Successful onboarding of 3 enterprise clients (Meridian Corp, TechVista, NovaStar)', style='List Bullet')

    # --- Section 3: Design Review ---
    doc.add_heading('3. Design and UX Review', level=2)
    doc.add_paragraph(
        'Elena Rodriguez shared the design team\'s accomplishments. The redesigned onboarding '
        'flow has reduced drop-off rates by 31%, and the new component library (internally '
        'called "Aurora") is now being adopted across all product teams. User satisfaction '
        'scores from the latest NPS survey improved from 42 to 58.'
    )
    doc.add_paragraph(
        'Upcoming design initiatives for Q2 include a complete overhaul of the settings page, '
        'accessibility compliance audit for WCAG 2.2 AA standards, and the introduction of a '
        'dark mode theme requested by 67% of surveyed users.'
    )

    # --- Section 4: QA Report ---
    doc.add_heading('4. Quality Assurance Report', level=2)
    doc.add_paragraph(
        'David Kim reported on testing outcomes. The QA team processed 1,247 test cases during '
        'Q1, identifying 89 bugs of which 73 have been resolved. Critical defect count dropped '
        'from 12 in Q4 to 4 in Q1. The team introduced automated regression testing for the '
        'payment processing module, reducing manual testing effort by approximately 40 hours '
        'per release cycle.'
    )

    # --- Page Break for page 2 ---
    doc.add_page_break()

    # --- Section 5: Financial Overview ---
    doc.add_heading('5. Financial Overview', level=2)
    doc.add_paragraph(
        'James O\'Brien presented the Q1 financial summary. Total expenditure for the engineering '
        'division was $2.34M against a budget of $2.50M, representing a 6.4% underspend. '
        'Infrastructure costs decreased by 18% following the cloud optimization initiative '
        'completed in February. Revenue from enterprise licenses grew 15% quarter-over-quarter, '
        'reaching $4.7M.'
    )
    doc.add_paragraph(
        'The CFO recommended reallocating the $160K surplus toward the Q2 hiring plan, which '
        'targets 5 additional senior engineers and 2 UX researchers. The proposal was approved '
        'unanimously by all attendees.'
    )

    # --- Section 6: Marketing Update ---
    doc.add_heading('6. Marketing and Growth', level=2)
    doc.add_paragraph(
        'Priya Patel outlined marketing performance. The content marketing strategy generated '
        '12,400 qualified leads in Q1, surpassing the target of 10,000 by 24%. Social media '
        'engagement increased 45% on LinkedIn and 32% on Twitter. The webinar series attracted '
        'an average of 340 registrants per session with a 62% attendance rate.'
    )

    # --- Section 7: Action Items ---
    doc.add_heading('7. Action Items', level=2)
    actions = [
        ('Marcus Johnson', 'Finalize Q2 product roadmap and circulate by March 22', 'March 22, 2026'),
        ('Elena Rodriguez', 'Complete WCAG 2.2 AA accessibility audit report', 'April 5, 2026'),
        ('David Kim', 'Implement automated E2E testing for checkout flow', 'April 12, 2026'),
        ('Priya Patel', 'Prepare Q2 marketing budget proposal with ROI projections', 'March 28, 2026'),
        ('James O\'Brien', 'Process Q2 hiring requisitions for engineering and design', 'April 1, 2026'),
        ('Sarah Chen', 'Schedule individual team sync meetings for Q2 planning', 'March 20, 2026'),
    ]
    for owner, action, deadline in actions:
        para = doc.add_paragraph(style='List Bullet')
        run_owner = para.add_run(f'{owner}: ')
        run_owner.bold = True
        para.add_run(f'{action} (Due: {deadline})')

    # --- Closing ---
    doc.add_heading('8. Next Meeting', level=2)
    doc.add_paragraph(
        'The next quarterly review is scheduled for June 18, 2026, at 10:00 AM in Conference '
        'Room B. Sarah Chen will distribute the agenda one week in advance. The meeting was '
        'adjourned at 11:28 AM.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
