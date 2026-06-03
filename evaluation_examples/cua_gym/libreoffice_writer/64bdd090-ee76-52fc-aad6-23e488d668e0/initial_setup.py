"""
Initial Setup: Product page document without text frames callout section
Task ID: writer_obj_063
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'product_page'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Product title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('NovaTech ProSuite 2025')
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    title_para.paragraph_format.space_after = Pt(6)
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run('The All-in-One Business Productivity Platform')
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    subtitle_para.paragraph_format.space_after = Pt(12)

    # Section divider line (using paragraph border)
    divider = doc.add_paragraph()
    divider.paragraph_format.space_before = Pt(0)
    divider.paragraph_format.space_after = Pt(10)

    # Overview section heading
    overview_heading = doc.add_paragraph()
    oh_run = overview_heading.add_run('Product Overview')
    oh_run.bold = True
    oh_run.font.size = Pt(14)
    oh_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    overview_heading.paragraph_format.space_before = Pt(4)
    overview_heading.paragraph_format.space_after = Pt(6)

    # Overview paragraph
    overview_para = doc.add_paragraph()
    overview_run = overview_para.add_run(
        'NovaTech ProSuite 2025 is a comprehensive business productivity solution '
        'designed for modern enterprises. Built on a cloud-native architecture, it '
        'seamlessly integrates project management, customer relationship management, '
        'and advanced analytics into a single unified platform. Organizations of all '
        'sizes can streamline their workflows and achieve measurable efficiency gains '
        'from day one of deployment.'
    )
    overview_run.font.size = Pt(11)
    overview_para.paragraph_format.space_after = Pt(8)

    # Key benefits heading
    benefits_heading = doc.add_paragraph()
    bh_run = benefits_heading.add_run('Key Benefits')
    bh_run.bold = True
    bh_run.font.size = Pt(14)
    bh_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    benefits_heading.paragraph_format.space_before = Pt(4)
    benefits_heading.paragraph_format.space_after = Pt(6)

    # Benefits paragraph
    benefits_para = doc.add_paragraph()
    benefits_run = benefits_para.add_run(
        'Enterprises adopting NovaTech ProSuite 2025 report an average of 34% reduction '
        'in operational overhead within the first quarter. The platform\'s AI-powered '
        'recommendation engine identifies bottlenecks proactively, while the real-time '
        'collaboration suite enables distributed teams across multiple time zones to '
        'coordinate with zero latency issues. Security-conscious organizations benefit '
        'from ISO 27001 certified infrastructure and end-to-end encryption for all '
        'data in transit and at rest.'
    )
    benefits_run.font.size = Pt(11)
    benefits_para.paragraph_format.space_after = Pt(8)

    # Features intro paragraph (content that will appear above the callout section at Y~12cm)
    features_intro = doc.add_paragraph()
    fi_run = features_intro.add_run(
        'The platform ships with three powerful modules that address the most critical '
        'needs of today\'s businesses. Each module is independently scalable and can be '
        'licensed separately or as part of the complete ProSuite bundle:'
    )
    fi_run.font.size = Pt(11)
    features_intro.paragraph_format.space_after = Pt(10)

    # Note: The callout section (three text frames) will be inserted by the agent below this content.

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
