"""
Initial Setup: Classic Tomato Soup recipe with six plain-text instruction steps.
Task ID: writer_lec_003
Domain: libreoffice_writer

The document contains a recipe title and six instruction steps as plain
paragraphs (no list formatting).  The agent's job is to convert these
into a numbered list with lowercase Roman numerals.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # ----- Title -----
    title = doc.add_heading('Classic Tomato Soup', level=1)

    # ----- Intro paragraph -----
    intro = doc.add_paragraph(
        'This comforting homemade tomato soup is perfect for a chilly evening. '
        'Rich, creamy, and bursting with fresh tomato flavour, it pairs '
        'beautifully with crusty bread or a grilled cheese sandwich.'
    )

    # ----- Ingredients sub-heading -----
    doc.add_heading('Ingredients', level=2)
    ingredients = [
        '2 lbs ripe Roma tomatoes, halved',
        '1 medium yellow onion, diced',
        '3 cloves garlic, minced',
        '2 tablespoons extra-virgin olive oil',
        '2 cups vegetable broth',
        '1/2 cup heavy cream',
        '1 teaspoon dried basil',
        '1/2 teaspoon smoked paprika',
        'Salt and freshly ground black pepper to taste',
        'Fresh basil leaves for garnish',
    ]
    for item in ingredients:
        doc.add_paragraph(item, style='List Bullet')

    # ----- Instructions sub-heading -----
    doc.add_heading('Instructions', level=2)

    # Six instruction steps as PLAIN paragraphs (no list style)
    steps = [
        'Preheat the oven to 400 degrees Fahrenheit and line a baking sheet '
        'with parchment paper. Place the halved Roma tomatoes cut-side up on '
        'the sheet, drizzle with one tablespoon of olive oil, and season with '
        'salt and pepper. Roast for 25 to 30 minutes until the edges are '
        'slightly charred.',

        'While the tomatoes are roasting, heat the remaining tablespoon of '
        'olive oil in a large pot over medium heat. Add the diced onion and '
        'cook for about 5 minutes, stirring occasionally, until softened and '
        'translucent.',

        'Add the minced garlic to the pot and cook for another 30 seconds '
        'until fragrant. Be careful not to let the garlic burn.',

        'Transfer the roasted tomatoes and any accumulated juices into the '
        'pot. Pour in the vegetable broth, dried basil, and smoked paprika. '
        'Stir to combine and bring the mixture to a gentle simmer. Let it '
        'cook for 10 minutes to meld the flavours.',

        'Remove the pot from heat and use an immersion blender to puree the '
        'soup until smooth. Alternatively, carefully transfer the soup in '
        'batches to a countertop blender. Stir in the heavy cream and return '
        'to low heat for 2 to 3 minutes.',

        'Taste and adjust seasoning with salt and pepper as needed. Ladle '
        'the soup into bowls, garnish with fresh basil leaves, and serve '
        'immediately with warm crusty bread.',
    ]

    for step_text in steps:
        doc.add_paragraph(step_text)

    # ----- Notes section -----
    doc.add_heading('Notes', level=2)
    doc.add_paragraph(
        'For a dairy-free version, substitute the heavy cream with coconut '
        'cream. The soup can be stored in an airtight container in the '
        'refrigerator for up to 4 days or frozen for up to 3 months.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer for the GUI agent
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
