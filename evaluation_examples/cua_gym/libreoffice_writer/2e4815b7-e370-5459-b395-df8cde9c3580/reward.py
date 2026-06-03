"""
Reward Script: Change Management Plan - ERP System Migration
Task ID: writer_wf_070
Domain: libreoffice_writer
Scoring:
  Component 1: Title present and correct (0.10)
  Component 2: 7 Heading 1 sections with correct names (0.25)
  Component 3: Impact Analysis table - header + 5 dept rows, 3 cols (0.20)
  Component 4: Stakeholder Analysis table - header + 4 rows, 4 cols (0.15)
  Component 5: Communication Plan table - header + 4 rows, 4 cols (0.15)
  Component 6: Training Plan table - header + 3 rows, 4 cols (0.15)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_070'


def persist_app_state(domain):
    """Send Ctrl+S to save any unsaved LibreOffice edits."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraphs with their styles
    all_paras = [(p.style.name, p.text.strip()) for p in doc.paragraphs]

    # Component 1: Title present and correct (0.10 points)
    # The title should be "Change Management Plan - ERP System Migration"
    try:
        title_paras = [(s, t) for s, t in all_paras if s == 'Title' and t]
        if title_paras:
            title_text = title_paras[0][1].lower()
            if 'change management plan' in title_text and 'erp' in title_text:
                print(f"PASS: Component 1 — Title found: '{title_paras[0][1][:80]}' (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 1 — Title text doesn't match: '{title_paras[0][1][:80]}'")
        else:
            print("FAIL: Component 1 — No paragraph with 'Title' style found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 7 Heading 1 sections with correct names (0.25 points)
    # Expected sections: Change Overview, Impact Analysis, Stakeholder Analysis,
    # Communication Plan, Training Plan, Resistance Management, Success Metrics
    try:
        expected_sections = [
            'change overview',
            'impact analysis',
            'stakeholder analysis',
            'communication plan',
            'training plan',
            'resistance management',
            'success metrics',
        ]
        heading1_paras = [(s, t) for s, t in all_paras if s == 'Heading 1' and t]
        heading1_texts_lower = [t.lower() for _, t in heading1_paras]

        matched_sections = 0
        for expected in expected_sections:
            if any(expected in h for h in heading1_texts_lower):
                matched_sections += 1

        if matched_sections == 7:
            print(f"PASS: Component 2 — All 7 Heading 1 sections found (0.25 pts)")
            total_score += 0.25
        elif matched_sections >= 5:
            partial = round(0.25 * (matched_sections / 7), 2)
            print(f"PARTIAL: Component 2 — {matched_sections}/7 sections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {matched_sections}/7 Heading 1 sections found. "
                  f"Found headings: {[t for _, t in heading1_paras]}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Helper: find a table near a specific heading
    # We match tables by position relative to headings, or by header row content
    tables = doc.tables

    # Component 3: Impact Analysis table (0.20 points)
    # Expected: header + 5 department rows, 3 columns (Department, Impact Level, Key Changes)
    try:
        impact_table = None
        for table in tables:
            header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            if 'department' in header_cells and 'impact level' in header_cells:
                impact_table = table
                break

        if impact_table is not None:
            num_rows = len(impact_table.rows)
            num_cols = len(impact_table.columns)
            # We expect header + 5 data rows = 6 rows, and 3 columns
            pts = 0.0
            if num_cols >= 3:
                pts += 0.05
            if num_rows >= 6:  # header + 5 departments
                pts += 0.10
            # Check that data rows have content (not empty)
            data_rows_with_content = 0
            for ri in range(1, min(num_rows, 6)):
                if impact_table.rows[ri].cells[0].text.strip():
                    data_rows_with_content += 1
            if data_rows_with_content >= 5:
                pts += 0.05
            print(f"PASS: Component 3 — Impact Analysis table: {num_rows} rows, {num_cols} cols, "
                  f"{data_rows_with_content} filled dept rows ({pts} pts)")
            total_score += pts
        else:
            print("FAIL: Component 3 — No Impact Analysis table found (looking for 'Department' + 'Impact Level' headers)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Stakeholder Analysis table (0.15 points)
    # Expected: header + 4 rows, 4 columns (Stakeholder Group, Influence, Interest, Strategy)
    try:
        stakeholder_table = None
        for table in tables:
            header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            if 'stakeholder' in ' '.join(header_cells) and 'influence' in header_cells:
                stakeholder_table = table
                break

        if stakeholder_table is not None:
            num_rows = len(stakeholder_table.rows)
            num_cols = len(stakeholder_table.columns)
            pts = 0.0
            if num_cols >= 4:
                pts += 0.05
            if num_rows >= 5:  # header + 4 groups
                pts += 0.05
            data_rows_with_content = 0
            for ri in range(1, min(num_rows, 5)):
                if stakeholder_table.rows[ri].cells[0].text.strip():
                    data_rows_with_content += 1
            if data_rows_with_content >= 4:
                pts += 0.05
            print(f"PASS: Component 4 — Stakeholder table: {num_rows} rows, {num_cols} cols, "
                  f"{data_rows_with_content} filled rows ({pts} pts)")
            total_score += pts
        else:
            print("FAIL: Component 4 — No Stakeholder Analysis table found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Communication Plan table (0.15 points)
    # Expected: header + 4 rows, 4 columns (Audience, Message, Channel, Frequency)
    try:
        comms_table = None
        for table in tables:
            header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            if 'audience' in header_cells and 'channel' in header_cells:
                comms_table = table
                break

        if comms_table is not None:
            num_rows = len(comms_table.rows)
            num_cols = len(comms_table.columns)
            pts = 0.0
            if num_cols >= 4:
                pts += 0.05
            if num_rows >= 5:  # header + 4 rows
                pts += 0.05
            data_rows_with_content = 0
            for ri in range(1, min(num_rows, 5)):
                if comms_table.rows[ri].cells[0].text.strip():
                    data_rows_with_content += 1
            if data_rows_with_content >= 4:
                pts += 0.05
            print(f"PASS: Component 5 — Communication Plan table: {num_rows} rows, {num_cols} cols, "
                  f"{data_rows_with_content} filled rows ({pts} pts)")
            total_score += pts
        else:
            print("FAIL: Component 5 — No Communication Plan table found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Training Plan table (0.15 points)
    # Expected: header + 3 rows, 4 columns (Group, Training, Duration, Date)
    try:
        training_table = None
        for table in tables:
            header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            if 'training' in header_cells and 'duration' in header_cells:
                training_table = table
                break

        if training_table is not None:
            num_rows = len(training_table.rows)
            num_cols = len(training_table.columns)
            pts = 0.0
            if num_cols >= 4:
                pts += 0.05
            if num_rows >= 4:  # header + 3 groups
                pts += 0.05
            data_rows_with_content = 0
            for ri in range(1, min(num_rows, 4)):
                if training_table.rows[ri].cells[0].text.strip():
                    data_rows_with_content += 1
            if data_rows_with_content >= 3:
                pts += 0.05
            print(f"PASS: Component 6 — Training Plan table: {num_rows} rows, {num_cols} cols, "
                  f"{data_rows_with_content} filled rows ({pts} pts)")
            total_score += pts
        else:
            print("FAIL: Component 6 — No Training Plan table found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook — save any unsaved LibreOffice edits before verification
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
