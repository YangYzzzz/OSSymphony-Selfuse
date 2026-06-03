"""
Reward Script: Software release notes document verification
Task ID: writer_wf_089
Domain: libreoffice_writer
Scoring:
  C1  Title present (0.10)
  C2  Release date + build number (0.10)
  C3  7 section headings (0.15)
  C4  What's New — 5 bulleted features (0.10)
  C5  Improvements — 4 bulleted items (0.10)
  C6  Bug Fixes table — header + 6 data rows, 3 cols (0.15)
  C7  Known Issues — 3 bulleted items (0.05)
  C8  Breaking Changes with migration content (0.05)
  C9  System Requirements table — 4 data rows, 2 cols (0.10)
  C10 Upgrade Instructions — 5 numbered steps (0.10)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_089'


def verify_task(file_path):
    """Verify release notes document with progressive scoring. Returns 0.0-1.0."""
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    paragraphs = doc.paragraphs
    tables = doc.tables

    # Build helper structures
    headings = []
    for p in paragraphs:
        if p.style and p.style.name and p.style.name.startswith('Heading'):
            headings.append(p.text.strip())

    # Map sections: heading text -> list of paragraphs following it (until next heading)
    sections = {}
    current_heading = None
    for p in paragraphs:
        if p.style and p.style.name and p.style.name.startswith('Heading'):
            current_heading = p.text.strip()
            sections[current_heading] = []
        elif current_heading is not None:
            sections[current_heading].append(p)

    # --- Component 1: Title present (0.10) ---
    try:
        title_found = False
        for p in paragraphs:
            if p.style and p.style.name == 'Title':
                if 'release notes' in p.text.lower() and 'cloudsync' in p.text.lower():
                    title_found = True
                    break
        if title_found:
            print(f"PASS: C1 — Title found: '{p.text}' (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: C1 — No Title paragraph with 'Release Notes' and 'CloudSync'")
    except Exception as e:
        print(f"ERROR: C1 — {e}")

    # --- Component 2: Release date + build number (0.10) ---
    try:
        full_text = '\n'.join(p.text for p in paragraphs)
        has_date = any(kw in full_text.lower() for kw in ['release date', 'date:'])
        has_build = any(kw in full_text.lower() for kw in ['build number', 'build:'])
        date_score = 0.0
        if has_date:
            date_score += 0.05
        if has_build:
            date_score += 0.05
        if date_score > 0:
            print(f"PASS: C2 — Release date={has_date}, Build number={has_build} ({date_score} pts)")
            total_score += date_score
        else:
            print(f"FAIL: C2 — No release date or build number found")
    except Exception as e:
        print(f"ERROR: C2 — {e}")

    # --- Component 3: 7 section headings (0.15) ---
    try:
        expected_sections = [
            "what's new", "improvements", "bug fixes",
            "known issues", "breaking changes",
            "system requirements", "upgrade instructions"
        ]
        found_sections = []
        for exp in expected_sections:
            for h in headings:
                if exp in h.lower():
                    found_sections.append(exp)
                    break
        count = len(found_sections)
        if count >= 7:
            pts = 0.15
        elif count >= 5:
            pts = 0.10
        elif count >= 3:
            pts = 0.05
        else:
            pts = 0.0
        if pts > 0:
            print(f"PASS: C3 — {count}/7 section headings found ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: C3 — Only {count}/7 section headings found")
    except Exception as e:
        print(f"ERROR: C3 — {e}")

    # --- Component 4: What's New — 5 bulleted features (0.10) ---
    try:
        whats_new_bullets = 0
        for key, paras in sections.items():
            if "what's new" in key.lower() or "whats new" in key.lower():
                for p in paras:
                    if p.style and 'bullet' in p.style.name.lower() and p.text.strip():
                        whats_new_bullets += 1
                break
        if whats_new_bullets >= 5:
            pts = 0.10
        elif whats_new_bullets >= 3:
            pts = 0.05
        else:
            pts = 0.0
        if pts > 0:
            print(f"PASS: C4 — What's New has {whats_new_bullets} bullet items ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: C4 — What's New has {whats_new_bullets} bullet items (need >= 3)")
    except Exception as e:
        print(f"ERROR: C4 — {e}")

    # --- Component 5: Improvements — 4 bulleted items (0.10) ---
    try:
        improvements_bullets = 0
        for key, paras in sections.items():
            if 'improvement' in key.lower():
                for p in paras:
                    if p.style and 'bullet' in p.style.name.lower() and p.text.strip():
                        improvements_bullets += 1
                break
        if improvements_bullets >= 4:
            pts = 0.10
        elif improvements_bullets >= 2:
            pts = 0.05
        else:
            pts = 0.0
        if pts > 0:
            print(f"PASS: C5 — Improvements has {improvements_bullets} bullet items ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: C5 — Improvements has {improvements_bullets} bullet items (need >= 2)")
    except Exception as e:
        print(f"ERROR: C5 — {e}")

    # --- Component 6: Bug Fixes table — header + 6 data rows, 3 cols (0.15) ---
    try:
        bug_table_found = False
        bug_table_score = 0.0

        # Find the table that follows the Bug Fixes heading.
        # Strategy: look for a table whose first-row cells contain Issue ID / Description / Severity
        for t in tables:
            header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
            has_issue_col = any('issue' in c or 'id' in c for c in header_cells)
            has_desc_col = any('desc' in c for c in header_cells)
            has_sev_col = any('sever' in c for c in header_cells)
            if has_issue_col and has_desc_col and has_sev_col:
                bug_table_found = True
                num_cols = len(t.columns)
                num_data_rows = len(t.rows) - 1  # exclude header
                # Score: correct columns (0.05), >= 6 data rows (0.10)
                if num_cols >= 3:
                    bug_table_score += 0.05
                if num_data_rows >= 6:
                    bug_table_score += 0.10
                elif num_data_rows >= 3:
                    bug_table_score += 0.05
                break

        if bug_table_score > 0:
            print(f"PASS: C6 — Bug Fixes table found with {num_data_rows} data rows, {num_cols} cols ({bug_table_score} pts)")
            total_score += bug_table_score
        else:
            print(f"FAIL: C6 — Bug Fixes table not found or insufficient structure (found={bug_table_found})")
    except Exception as e:
        print(f"ERROR: C6 — {e}")

    # --- Component 7: Known Issues — 3 bulleted items (0.05) ---
    try:
        known_issues_bullets = 0
        for key, paras in sections.items():
            if 'known issue' in key.lower():
                for p in paras:
                    if p.style and 'bullet' in p.style.name.lower() and p.text.strip():
                        known_issues_bullets += 1
                break
        if known_issues_bullets >= 3:
            pts = 0.05
        else:
            pts = 0.0
        if pts > 0:
            print(f"PASS: C7 — Known Issues has {known_issues_bullets} bullet items ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: C7 — Known Issues has {known_issues_bullets} bullet items (need >= 3)")
    except Exception as e:
        print(f"ERROR: C7 — {e}")

    # --- Component 8: Breaking Changes with migration content (0.05) ---
    try:
        breaking_score = 0.0
        for key, paras in sections.items():
            if 'breaking' in key.lower():
                text_content = ' '.join(p.text for p in paras).lower()
                # Must have some migration-related content
                if any(kw in text_content for kw in ['migrat', 'updat', 'replac', 'url', 'api']):
                    breaking_score = 0.05
                break
        if breaking_score > 0:
            print(f"PASS: C8 — Breaking Changes section with migration content ({breaking_score} pts)")
            total_score += breaking_score
        else:
            print(f"FAIL: C8 — Breaking Changes missing or no migration content")
    except Exception as e:
        print(f"ERROR: C8 — {e}")

    # --- Component 9: System Requirements table — 4 data rows, 2 cols (0.10) ---
    try:
        req_table_found = False
        req_table_score = 0.0

        for t in tables:
            header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
            has_comp = any('component' in c for c in header_cells)
            has_req = any('requirement' in c for c in header_cells)
            if has_comp and has_req:
                req_table_found = True
                num_data_rows = len(t.rows) - 1
                num_cols = len(t.columns)
                if num_cols >= 2:
                    req_table_score += 0.03
                if num_data_rows >= 4:
                    req_table_score += 0.07
                elif num_data_rows >= 2:
                    req_table_score += 0.04
                break

        if req_table_score > 0:
            print(f"PASS: C9 — System Requirements table: {num_data_rows} data rows, {num_cols} cols ({req_table_score} pts)")
            total_score += req_table_score
        else:
            print(f"FAIL: C9 — System Requirements table not found or insufficient (found={req_table_found})")
    except Exception as e:
        print(f"ERROR: C9 — {e}")

    # --- Component 10: Upgrade Instructions — 5 numbered steps (0.10) ---
    try:
        upgrade_numbered = 0
        for key, paras in sections.items():
            if 'upgrade' in key.lower() and 'instruct' in key.lower():
                for p in paras:
                    if p.style and 'number' in p.style.name.lower() and p.text.strip():
                        upgrade_numbered += 1
                break
        if upgrade_numbered >= 5:
            pts = 0.10
        elif upgrade_numbered >= 3:
            pts = 0.05
        else:
            pts = 0.0
        if pts > 0:
            print(f"PASS: C10 — Upgrade Instructions has {upgrade_numbered} numbered steps ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: C10 — Upgrade Instructions has {upgrade_numbered} numbered steps (need >= 3)")
    except Exception as e:
        print(f"ERROR: C10 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
