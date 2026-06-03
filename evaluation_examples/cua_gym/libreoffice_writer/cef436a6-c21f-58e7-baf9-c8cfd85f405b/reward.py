"""
Reward Script: HR Compliance Audit Checklist
Task ID: writer_hr_079
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.25): Summary dashboard table with 6 categories
  - Component 2 (0.35): 6 section tables with correct item counts (45 total)
  - Component 3 (0.20): All section tables have 6-column structure
  - Component 4 (0.20): Section headings for all 6 categories
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_079'

# Expected section names and their required item counts (data rows, excluding header)
EXPECTED_SECTIONS = {
    'Equal Employment Opportunity': 8,
    'Wage & Hour': 10,
    'Benefits Compliance': 7,
    'Safety & Health': 9,
    'Record Keeping': 6,
    'Immigration': 5,
}

EXPECTED_COLUMNS = ['Requirement', 'Compliant', 'Evidence Location', 'Last Reviewed', 'Reviewer', 'Notes']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather all paragraphs and tables
    paragraphs = doc.paragraphs
    tables = doc.tables

    # --- Component 1: Summary dashboard table at top with 6 categories (0.25 pts) ---
    # The summary dashboard should be the first table, with category rows matching
    # the 6 sections and columns like Category, Compliance Rate, etc.
    try:
        if len(tables) >= 1:
            dashboard = tables[0]
            # Check that it has at least 7 rows (1 header + 6 categories)
            dashboard_rows = len(dashboard.rows)
            # Check header row has "Category" or similar
            header_cells = [c.text.strip().lower() for c in dashboard.rows[0].cells]
            has_category_header = any('category' in h for h in header_cells)
            has_compliance_header = any('compliance' in h or 'rate' in h for h in header_cells)

            # Check that all 6 section names appear in the dashboard data rows
            dashboard_categories = set()
            for ri in range(1, dashboard_rows):
                first_cell = dashboard.rows[ri].cells[0].text.strip()
                for section_name in EXPECTED_SECTIONS:
                    if section_name.lower() in first_cell.lower():
                        dashboard_categories.add(section_name)

            categories_found = len(dashboard_categories)

            if has_category_header and has_compliance_header and categories_found == 6:
                print(f"PASS: Component 1 — Summary dashboard has all 6 categories with proper headers (0.25 pts)")
                total_score += 0.25
            elif categories_found >= 4:
                partial = 0.15
                print(f"PARTIAL: Component 1 — Dashboard found {categories_found}/6 categories ({partial} pts)")
                total_score += partial
            elif categories_found >= 1:
                partial = 0.05
                print(f"PARTIAL: Component 1 — Dashboard found {categories_found}/6 categories ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 1 — No valid dashboard table found (categories={categories_found}, cat_header={has_category_header}, compliance_header={has_compliance_header})")
        else:
            print("FAIL: Component 1 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # --- Component 2: 6 section tables with correct item counts totaling 45 (0.35 pts) ---
    # Section tables are tables[1] through tables[6] (after the dashboard).
    # Each must have the right number of data rows (header row + N items).
    try:
        section_tables = tables[1:]  # skip dashboard
        sections_matched = 0
        total_items = 0

        if len(section_tables) >= 6:
            expected_counts = list(EXPECTED_SECTIONS.values())  # [8, 10, 7, 9, 6, 5]
            for idx, expected_count in enumerate(expected_counts):
                if idx < len(section_tables):
                    t = section_tables[idx]
                    data_rows = len(t.rows) - 1  # subtract header
                    total_items += data_rows
                    if data_rows == expected_count:
                        sections_matched += 1
                        print(f"  Section table {idx+1}: {data_rows} items (expected {expected_count}) — OK")
                    else:
                        print(f"  Section table {idx+1}: {data_rows} items (expected {expected_count}) — MISMATCH")

            if sections_matched == 6 and total_items == 45:
                print(f"PASS: Component 2 — All 6 section tables have correct item counts, total 45 items (0.35 pts)")
                total_score += 0.35
            elif sections_matched >= 4:
                partial = round(0.35 * sections_matched / 6, 2)
                print(f"PARTIAL: Component 2 — {sections_matched}/6 tables have correct counts ({partial} pts)")
                total_score += partial
            elif sections_matched >= 1:
                partial = round(0.35 * sections_matched / 6, 2)
                print(f"PARTIAL: Component 2 — {sections_matched}/6 tables have correct counts ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — No section tables have correct item counts (matched={sections_matched}, total_items={total_items})")
        else:
            print(f"FAIL: Component 2 — Only {len(section_tables)} section tables found (expected 6)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # --- Component 3: All section tables have 6-column structure (0.20 pts) ---
    # Each section table header row must contain the 6 expected columns.
    try:
        section_tables = tables[1:]
        tables_with_correct_cols = 0

        for idx in range(min(6, len(section_tables))):
            t = section_tables[idx]
            if len(t.rows) > 0:
                header_texts = [c.text.strip() for c in t.rows[0].cells]
                # Check each expected column name is present
                matched_cols = 0
                for exp_col in EXPECTED_COLUMNS:
                    if any(exp_col.lower() in h.lower() for h in header_texts):
                        matched_cols += 1
                if matched_cols == 6:
                    tables_with_correct_cols += 1
                else:
                    print(f"  Section table {idx+1} headers: {header_texts} — only {matched_cols}/6 columns matched")

        if tables_with_correct_cols == 6:
            print(f"PASS: Component 3 — All 6 section tables have correct 6-column structure (0.20 pts)")
            total_score += 0.20
        elif tables_with_correct_cols >= 3:
            partial = round(0.20 * tables_with_correct_cols / 6, 2)
            print(f"PARTIAL: Component 3 — {tables_with_correct_cols}/6 tables have correct columns ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {tables_with_correct_cols}/6 tables have correct column structure")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # --- Component 4: Section headings for all 6 categories (0.20 pts) ---
    # The document must have heading paragraphs matching the 6 section names.
    try:
        heading_texts = []
        for p in paragraphs:
            if p.style and 'Heading' in p.style.name:
                heading_texts.append(p.text.strip())

        sections_found = 0
        for section_name in EXPECTED_SECTIONS:
            if any(section_name.lower() in h.lower() for h in heading_texts):
                sections_found += 1
            else:
                print(f"  Missing heading for: {section_name}")

        if sections_found == 6:
            print(f"PASS: Component 4 — All 6 section headings present (0.20 pts)")
            total_score += 0.20
        elif sections_found >= 3:
            partial = round(0.20 * sections_found / 6, 2)
            print(f"PARTIAL: Component 4 — {sections_found}/6 section headings found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {sections_found}/6 section headings found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
