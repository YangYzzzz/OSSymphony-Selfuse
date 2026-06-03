"""
Initial Setup: Create a technical documentation document with a plain NOTE paragraph.
Task ID: writer_tech_036
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_036'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('CloudSync API Reference Guide', level=0)

    # --- Introduction ---
    intro = doc.add_heading('1. Introduction', level=1)
    p = doc.add_paragraph(
        'CloudSync is a distributed file synchronization platform designed for '
        'enterprise environments. It provides real-time replication across data '
        'centers with configurable consistency guarantees and conflict resolution '
        'policies. This guide covers the core REST API endpoints, authentication '
        'mechanisms, and webhook integration patterns.'
    )

    # --- Authentication section ---
    auth_heading = doc.add_heading('2. Authentication', level=1)
    p = doc.add_paragraph(
        'All API requests must include a valid bearer token in the Authorization '
        'header. Tokens are obtained through the OAuth 2.0 client credentials flow '
        'using your application\'s client_id and client_secret. Tokens expire after '
        '3600 seconds by default.'
    )

    p = doc.add_paragraph(
        'To request a token, send a POST request to /auth/token with the following '
        'parameters:'
    )

    # Token parameters table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['Parameter', 'Type', 'Description']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data = [
        ['client_id', 'string', 'Your application client identifier'],
        ['client_secret', 'string', 'Your application secret key'],
        ['grant_type', 'string', 'Must be "client_credentials"'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')  # spacer

    # --- Sync Operations section ---
    sync_heading = doc.add_heading('3. Sync Operations', level=1)
    p = doc.add_paragraph(
        'The /sync endpoint manages bidirectional file synchronization between '
        'registered nodes. Each sync operation creates a transaction that can be '
        'monitored through the /status endpoint. Sync operations support three '
        'conflict resolution strategies: last-writer-wins, manual-merge, and '
        'version-branching.'
    )

    # Endpoint details table
    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'
    headers2 = ['Endpoint', 'Method', 'Auth Required', 'Rate Limit']
    for i, h in enumerate(headers2):
        cell = table2.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    endpoints = [
        ['/sync/start', 'POST', 'Yes', '100/min'],
        ['/sync/status/{id}', 'GET', 'Yes', '500/min'],
        ['/sync/cancel/{id}', 'DELETE', 'Yes', '50/min'],
        ['/sync/history', 'GET', 'Yes', '200/min'],
    ]
    for r, row_data in enumerate(endpoints, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    doc.add_paragraph('')  # spacer

    # --- Permissions section with the NOTE paragraph ---
    perm_heading = doc.add_heading('4. Permissions and Access Control', level=1)
    p = doc.add_paragraph(
        'CloudSync implements role-based access control (RBAC) with four predefined '
        'roles: Viewer, Editor, Manager, and Administrator. Each role grants '
        'progressively broader permissions over sync configurations, user management, '
        'and audit log access.'
    )

    # THE NOTE PARAGRAPH - plain formatting, no admonition styling
    note_para = doc.add_paragraph(
        'NOTE: This feature requires admin privileges. Only users with the '
        'Administrator role can modify access control lists, create custom roles, '
        'or configure cross-tenant sharing policies. Attempts to access these '
        'endpoints without sufficient permissions will return HTTP 403.'
    )

    p = doc.add_paragraph(
        'To assign roles, use the /users/{id}/roles endpoint with a PATCH request '
        'containing the desired role identifiers. Role changes take effect '
        'immediately and are logged in the audit trail.'
    )

    # --- Webhooks section ---
    wh_heading = doc.add_heading('5. Webhook Integration', level=1)
    p = doc.add_paragraph(
        'CloudSync supports webhook notifications for sync events, error conditions, '
        'and quota thresholds. Webhooks are configured per-organization and can be '
        'filtered by event type. All webhook payloads are signed using HMAC-SHA256 '
        'with your webhook secret for verification.'
    )

    p = doc.add_paragraph(
        'Supported webhook events include: sync.started, sync.completed, '
        'sync.failed, sync.conflict, quota.warning, and quota.exceeded. Each '
        'payload includes a timestamp, event type, affected resource IDs, and '
        'relevant metadata.'
    )

    # --- Error Handling section ---
    err_heading = doc.add_heading('6. Error Handling', level=1)
    p = doc.add_paragraph(
        'All API errors follow RFC 7807 Problem Details format. The response body '
        'includes a type URI, title, status code, detail message, and an optional '
        'instance identifier for support reference. Common error codes include '
        'SYNC_IN_PROGRESS (409), QUOTA_EXCEEDED (429), and NODE_UNREACHABLE (503).'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
