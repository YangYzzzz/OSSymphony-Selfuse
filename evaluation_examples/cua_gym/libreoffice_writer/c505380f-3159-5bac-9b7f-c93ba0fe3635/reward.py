"""
Reward Script: Insert OLE object linking to test_results.xlsx
Task ID: writer_tech_062
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): OLE object with Excel ProgID exists in the document
  Component 2 (0.3): Static table removed (0 tables in document)
  Component 3 (0.3): OLE object is located in the Test Results section
"""

import os
import time


WORKDIR = '/home/user'
TASK_ID = 'writer_tech_062'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from lxml import etree
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_o = 'urn:schemas-microsoft-com:office:office'
    ns = {'w': ns_w, 'o': ns_o}

    # Component 1: OLE object with Excel ProgID exists (0.4 points)
    # This checks that an OLE object referencing an Excel spreadsheet is embedded.
    # Initial file has 0 OLE objects; golden has 1 with ProgID=Excel.Sheet.12
    try:
        ole_objects = doc.element.body.findall('.//w:object', ns)
        excel_ole_found = False
        for obj in ole_objects:
            ole_elems = obj.findall('.//o:OLEObject', ns)
            for ole in ole_elems:
                prog_id = ole.get('ProgID')
                if prog_id and 'Excel' in prog_id:
                    excel_ole_found = True
                    print(f"PASS: Component 1 — Excel OLE object found (ProgID={prog_id}) (0.4 pts)")
                    break
            if excel_ole_found:
                break

        if excel_ole_found:
            total_score += 0.4
        else:
            # Also check for package relationships pointing to xlsx as fallback
            has_xlsx_embed = False
            for rel in doc.part.rels.values():
                if 'package' in rel.reltype.lower() and '.xlsx' in str(rel.target_ref).lower():
                    has_xlsx_embed = True
                    break
            if has_xlsx_embed:
                print(f"PASS: Component 1 — xlsx package relationship found (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 1 — No Excel OLE object found. Found {len(ole_objects)} w:object elements.")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Static table removed (0.3 points)
    # Initial file has 1 table (the static test results). Golden should have 0 tables.
    try:
        num_tables = len(doc.tables)
        if num_tables == 0:
            print(f"PASS: Component 2 — Static table removed, 0 tables in document (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Expected 0 tables (static table removed), found {num_tables}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: OLE object is in the Test Results section (0.3 points)
    # The OLE object should be located between "3. Test Results" and "4. Failure Analysis" headings
    try:
        test_results_idx = None
        failure_analysis_idx = None
        ole_para_idx = None

        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if 'Test Results' in text and p.style and 'Heading' in p.style.name:
                test_results_idx = i
            if 'Failure Analysis' in text and p.style and 'Heading' in p.style.name:
                failure_analysis_idx = i
            # Check if this paragraph contains an OLE object
            objs = p._element.findall('.//w:object', ns)
            if len(objs) > 0:
                ole_para_idx = i

        if ole_para_idx is not None and test_results_idx is not None:
            # OLE should be after Test Results heading
            in_section = ole_para_idx > test_results_idx
            if failure_analysis_idx is not None:
                in_section = in_section and ole_para_idx < failure_analysis_idx

            if in_section:
                print(f"PASS: Component 3 — OLE object at paragraph {ole_para_idx}, "
                      f"between Test Results ({test_results_idx}) and Failure Analysis ({failure_analysis_idx}) (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 3 — OLE object at paragraph {ole_para_idx}, "
                      f"not in Test Results section (heading at {test_results_idx}, next section at {failure_analysis_idx})")
        elif ole_para_idx is None:
            print(f"FAIL: Component 3 — No OLE object found in any paragraph")
        else:
            print(f"FAIL: Component 3 — Could not locate Test Results section heading")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
