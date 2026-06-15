"""
Initial Setup: Writer document with static test results table and companion spreadsheet.
Task ID: writer_tech_062
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_062'
OUTPUT_DOCX = f'{WORKDIR}/{TASK_ID}.docx'
OUTPUT_XLSX = f'{WORKDIR}/test_results.xlsx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_spreadsheet():
    """Create the test_results.xlsx spreadsheet that will be linked via OLE."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Test Results"

    # Header styling
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    headers = ["Test Case ID", "Module", "Description", "Status", "Execution Time (ms)", "Priority"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    # Realistic test result data
    data = [
        ["TC-001", "Authentication", "Login with valid credentials", "Pass", 245, "Critical"],
        ["TC-002", "Authentication", "Login with expired token", "Pass", 312, "High"],
        ["TC-003", "Authentication", "MFA verification flow", "Fail", 1580, "Critical"],
        ["TC-004", "User Profile", "Update display name", "Pass", 189, "Medium"],
        ["TC-005", "User Profile", "Upload avatar image (PNG)", "Pass", 876, "Low"],
        ["TC-006", "User Profile", "Change email with verification", "Pass", 2340, "High"],
        ["TC-007", "API Gateway", "Rate limiting enforcement", "Pass", 156, "Critical"],
        ["TC-008", "API Gateway", "JWT token validation", "Fail", 423, "Critical"],
        ["TC-009", "Data Export", "CSV export with 10k records", "Pass", 4521, "Medium"],
        ["TC-010", "Data Export", "PDF report generation", "Pass", 3210, "Medium"],
        ["TC-011", "Notifications", "Email dispatch queue processing", "Pass", 890, "High"],
        ["TC-012", "Notifications", "Push notification delivery", "Fail", 1245, "High"],
        ["TC-013", "Search", "Full-text index rebuild", "Pass", 7832, "Low"],
        ["TC-014", "Search", "Fuzzy matching accuracy", "Pass", 534, "Medium"],
        ["TC-015", "Payments", "Stripe webhook processing", "Pass", 678, "Critical"],
    ]

    pass_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
    fail_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")

    for r, row_data in enumerate(data, 2):
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.border = thin_border
            if c == 4:  # Status column
                cell.fill = pass_fill if val == "Pass" else fail_fill
                cell.alignment = Alignment(horizontal="center")
            if c == 5:  # Execution time
                cell.alignment = Alignment(horizontal="right")

    # Column widths
    ws.column_dimensions['A'].width = 14
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 36
    ws.column_dimensions['D'].width = 10
    ws.column_dimensions['E'].width = 22
    ws.column_dimensions['F'].width = 12

    wb.save(OUTPUT_XLSX)
    print(f'Spreadsheet created: {OUTPUT_XLSX}')


def create_document():
    """Create the Writer document with a static table of test results."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading("Platform Integration Test Report", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Metadata paragraph ---
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run("Version 2.4.1  |  Sprint 18  |  March 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph("")  # spacer

    # --- 1. Executive Summary ---
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This report summarizes the integration test results for the Platform v2.4.1 release "
        "candidate. Testing was conducted over a three-day window from March 24-26, 2026, covering "
        "all critical and high-priority test cases across five core modules: Authentication, "
        "User Profile, API Gateway, Data Export, and Notifications."
    )
    doc.add_paragraph(
        "Overall pass rate stands at 80% (12 of 15 test cases passed). Three failures were "
        "identified in MFA verification, JWT token validation, and push notification delivery. "
        "The Authentication and API Gateway failures are classified as release blockers requiring "
        "resolution before deployment to production."
    )

    # --- 2. Test Environment ---
    doc.add_heading("2. Test Environment", level=1)
    env_items = [
        "Operating System: Ubuntu 22.04 LTS (staging cluster)",
        "Database: PostgreSQL 15.2 with read replicas",
        "Application Server: Node.js 20.11 (3 instances behind HAProxy)",
        "Cache Layer: Redis 7.2 cluster (3 nodes)",
        "External Services: Stripe API (sandbox), SendGrid (test mode), Firebase Cloud Messaging (dev project)",
    ]
    for item in env_items:
        doc.add_paragraph(item, style="List Bullet")

    # --- 3. Test Results ---
    doc.add_heading("3. Test Results", level=1)
    doc.add_paragraph(
        "The following table provides a detailed breakdown of each test case, including the "
        "module under test, a brief description, execution status, measured execution time, "
        "and assigned priority level."
    )

    # Static table with same data as the spreadsheet
    table = doc.add_table(rows=16, cols=6)
    table.style = "Table Grid"

    headers = ["Test Case ID", "Module", "Description", "Status", "Exec Time (ms)", "Priority"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Set cell shading via XML
        shading = cell._element.get_or_add_tcPr().makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '2F5496',
        })
        cell._element.get_or_add_tcPr().append(shading)

    data = [
        ["TC-001", "Authentication", "Login with valid credentials", "Pass", "245", "Critical"],
        ["TC-002", "Authentication", "Login with expired token", "Pass", "312", "High"],
        ["TC-003", "Authentication", "MFA verification flow", "Fail", "1580", "Critical"],
        ["TC-004", "User Profile", "Update display name", "Pass", "189", "Medium"],
        ["TC-005", "User Profile", "Upload avatar image (PNG)", "Pass", "876", "Low"],
        ["TC-006", "User Profile", "Change email with verification", "Pass", "2340", "High"],
        ["TC-007", "API Gateway", "Rate limiting enforcement", "Pass", "156", "Critical"],
        ["TC-008", "API Gateway", "JWT token validation", "Fail", "423", "Critical"],
        ["TC-009", "Data Export", "CSV export with 10k records", "Pass", "4521", "Medium"],
        ["TC-010", "Data Export", "PDF report generation", "Pass", "3210", "Medium"],
        ["TC-011", "Notifications", "Email dispatch queue processing", "Pass", "890", "High"],
        ["TC-012", "Notifications", "Push notification delivery", "Fail", "1245", "High"],
        ["TC-013", "Search", "Full-text index rebuild", "Pass", "7832", "Low"],
        ["TC-014", "Search", "Fuzzy matching accuracy", "Pass", "534", "Medium"],
        ["TC-015", "Payments", "Stripe webhook processing", "Pass", "678", "Critical"],
    ]

    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)

    # --- 4. Failure Analysis ---
    doc.add_paragraph("")
    doc.add_heading("4. Failure Analysis", level=1)

    doc.add_heading("4.1 TC-003: MFA Verification Flow", level=2)
    doc.add_paragraph(
        "The MFA verification timed out after 1580ms when the TOTP service returned a 503 "
        "status. Root cause: the TOTP microservice connection pool was exhausted under concurrent "
        "load. Recommended fix: increase pool size from 10 to 25 connections and add circuit "
        "breaker with 2-second timeout."
    )

    doc.add_heading("4.2 TC-008: JWT Token Validation", level=2)
    doc.add_paragraph(
        "JWT validation failed for tokens issued with the RS256 algorithm when the API Gateway "
        "attempted to verify using the rotated public key. The key rotation job ran at 02:00 UTC "
        "but the gateway cache TTL was set to 24 hours. Fix: reduce cache TTL to 1 hour and "
        "implement key ID (kid) header matching."
    )

    doc.add_heading("4.3 TC-012: Push Notification Delivery", level=2)
    doc.add_paragraph(
        "Firebase Cloud Messaging returned a 401 Unauthorized error. The service account "
        "credentials in the staging environment had expired on March 20, 2026. This is an "
        "environment configuration issue, not a code defect. Fix: rotate FCM credentials and "
        "add monitoring for credential expiry."
    )

    # --- 5. Recommendations ---
    doc.add_heading("5. Recommendations", level=1)
    recommendations = [
        "Resolve TC-003 and TC-008 before production deployment (release blockers).",
        "Rotate FCM credentials in all non-production environments (TC-012).",
        "Add automated credential expiry monitoring to the CI/CD pipeline.",
        "Consider increasing test coverage for the Search module (currently at 67%).",
        "Schedule regression testing for the Authentication module after connection pool changes.",
    ]
    for i, rec in enumerate(recommendations, 1):
        doc.add_paragraph(f"{i}. {rec}")

    doc.save(OUTPUT_DOCX)
    print(f'Document created: {OUTPUT_DOCX}')


if __name__ == '__main__':
    create_spreadsheet()
    create_document()

    # GUI-ready: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT_DOCX}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')
