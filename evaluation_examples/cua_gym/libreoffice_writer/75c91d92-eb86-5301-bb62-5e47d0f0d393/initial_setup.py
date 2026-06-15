"""
Initial Setup: Replace all line breaks (Shift+Enter) with paragraph breaks (Enter)
Task ID: writer_frd_035
Domain: libreoffice_writer

Creates a document with 35 manual/soft line breaks (<w:br/>) simulating
web content pasted into Writer, where lines are separated by soft returns
instead of paragraph marks.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_035'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_soft_break(run):
    """Add a soft line break (<w:br/>) to the given run."""
    br = run._element.makeelement(qn('w:br'), {})
    run._element.append(br)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title paragraph (normal paragraph break after this)
    title = doc.add_heading("Sustainable Urban Development: A Comprehensive Overview", level=1)

    # We'll create several paragraphs, each containing multiple lines
    # joined by soft line breaks (Shift+Enter). Total: 35 soft breaks.
    # The content simulates pasted web content.

    # --- Section 1: Introduction (7 soft breaks) ---
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(6)
    lines_1 = [
        "Urban areas are home to more than 55% of the world's population, a proportion",
        "expected to increase to 68% by 2050. The challenges of managing urban growth",
        "include providing adequate housing, transportation, and public services to an",
        "ever-expanding population. Sustainable urban development seeks to address",
        "these challenges through integrated planning approaches that balance economic",
        "growth, social equity, and environmental protection. This overview examines",
        "key strategies being implemented in cities around the world to create more",
        "livable and resilient urban environments.",
    ]
    # 7 soft breaks between 8 lines
    for i, line in enumerate(lines_1):
        run = p1.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if i < len(lines_1) - 1:
            add_soft_break(run)

    # --- Section 2: Green Infrastructure (6 soft breaks) ---
    h2 = doc.add_heading("Green Infrastructure and Urban Ecology", level=2)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    lines_2 = [
        "Green infrastructure encompasses a network of natural and semi-natural areas",
        "designed to deliver ecosystem services within urban settings. Parks, green roofs,",
        "urban forests, and bioswales help manage stormwater runoff, reduce the urban",
        "heat island effect, and improve air quality. Cities like Singapore, Copenhagen,",
        "and Melbourne have invested heavily in green corridor networks that connect",
        "fragmented habitats. Research from the Barcelona Institute for Global Health",
        "shows that residents living near green spaces report 25% lower stress levels.",
    ]
    for i, line in enumerate(lines_2):
        run = p2.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if i < len(lines_2) - 1:
            add_soft_break(run)

    # --- Section 3: Transportation (5 soft breaks) ---
    h3 = doc.add_heading("Sustainable Transportation Systems", level=2)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    lines_3 = [
        "Effective public transit reduces carbon emissions by an estimated 37 million",
        "metric tons annually in the United States alone. Electric bus fleets are being",
        "deployed in Shenzhen (16,359 buses), Santiago, and Amsterdam. Bicycle-sharing",
        "programs have expanded to over 1,800 cities worldwide, with Paris's Velib",
        "system logging more than 300,000 trips per day during peak months. Pedestrian",
        "zones in city centers, such as those in Vienna and Ghent, have demonstrated",
    ]
    for i, line in enumerate(lines_3):
        run = p3.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if i < len(lines_3) - 1:
            add_soft_break(run)

    # --- Section 4: Energy Efficiency (5 soft breaks) ---
    h4 = doc.add_heading("Building Energy Efficiency Standards", level=2)

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(6)
    lines_4 = [
        "Buildings account for approximately 40% of global energy consumption and",
        "33% of greenhouse gas emissions. Passive house standards, originating in",
        "Germany in 1991, reduce heating energy by up to 90% compared to conventional",
        "construction. The European Union's Energy Performance of Buildings Directive",
        "requires all new buildings to be nearly zero-energy by 2021. Retrofit programs",
        "in New York City target a 30% reduction in building emissions by 2030.",
    ]
    for i, line in enumerate(lines_4):
        run = p4.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if i < len(lines_4) - 1:
            add_soft_break(run)

    # --- Section 5: Waste Management (6 soft breaks) ---
    h5 = doc.add_heading("Circular Economy and Waste Reduction", level=2)

    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(6)
    lines_5 = [
        "The circular economy model aims to eliminate waste through better design,",
        "reuse, and recycling of materials. San Francisco diverts 80% of its waste",
        "from landfills through comprehensive recycling and composting mandates.",
        "Extended producer responsibility programs in South Korea have achieved a",
        "packaging recycling rate of 64%, compared to the global average of 14%.",
        "Zero-waste grocery stores and refill stations are emerging in cities across",
        "Europe and North America, reducing single-use packaging by an estimated 70%.",
    ]
    for i, line in enumerate(lines_5):
        run = p5.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if i < len(lines_5) - 1:
            add_soft_break(run)

    # --- Section 6: Community Engagement (6 soft breaks) ---
    h6 = doc.add_heading("Community Engagement and Social Equity", level=2)

    p6 = doc.add_paragraph()
    p6.paragraph_format.space_after = Pt(6)
    lines_6 = [
        "Participatory budgeting, pioneered in Porto Alegre, Brazil in 1989, allows",
        "residents to directly decide how public funds are allocated. Over 7,000 cities",
        "worldwide now use some form of participatory governance. Digital platforms",
        "like Decidim enable large-scale citizen input on urban planning decisions.",
        "Equity-focused development ensures that sustainable improvements benefit all",
        "residents, particularly marginalized communities disproportionately affected",
        "by pollution, flooding, and inadequate infrastructure investments.",
    ]
    for i, line in enumerate(lines_6):
        run = p6.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if i < len(lines_6) - 1:
            add_soft_break(run)

    # Total soft breaks: 7 + 6 + 5 + 5 + 6 + 6 = 35

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Count soft breaks for verification
    doc2 = Document(OUTPUT)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    br_count = 0
    for para in doc2.paragraphs:
        for run in para.runs:
            for br in run.element.findall('.//w:br', ns):
                br_type = br.get(qn('w:type'))
                if br_type is None or br_type == 'textWrapping':
                    br_count += 1
    print(f'Soft line breaks in document: {br_count}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
