"""
Reward Script: Replace all line breaks (Shift+Enter) with paragraph breaks (Enter)
Task ID: writer_frd_035
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): All soft line breaks removed (count == 0)
  Component 2 (0.3): Paragraph count increased (>= 40, from original 12 + 35 breaks)
  Component 3 (0.2): Text content preserved (no data loss)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_035'


def persist_app_state(domain):
    """Save any unsaved GUI state before verification."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        import time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def count_line_breaks(doc):
    """Count soft line breaks (<w:br/> without type attribute) in the document."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            for br in run.element.findall('.//w:br', ns):
                br_type = br.attrib.get(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type',
                    None
                )
                # No type attribute = soft line break (Shift+Enter)
                # type="page" = page break, type="column" = column break
                if br_type is None:
                    count += 1
    return count


def get_full_text(doc):
    """Extract all text from paragraphs, normalizing whitespace."""
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # Join and normalize whitespace for comparison
    return re.sub(r'\s+', ' ', ' '.join(parts)).strip()


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All soft line breaks removed (0.5 points)
    # Initial has 35 line breaks; golden should have 0.
    # This is THE core task change.
    try:
        lb_count = count_line_breaks(doc)
        if lb_count == 0:
            print(f"PASS: Component 1 — No soft line breaks found (0.5 pts)")
            total_score += 0.5
        else:
            # Partial credit: proportional to how many were removed (out of 35)
            removed = max(0, 35 - lb_count)
            partial = 0.5 * (removed / 35)
            total_score += partial
            print(f"FAIL: Component 1 — Found {lb_count} soft line breaks remaining "
                  f"(expected 0). Partial: {partial:.2f} pts")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Paragraph count increased (0.3 points)
    # Initial: 12 paragraphs. After converting 35 line breaks to paragraph breaks,
    # we expect approximately 47 paragraphs (12 + 35).
    # Accept >= 40 as reasonable (some tolerance for formatting variations).
    try:
        para_count = len(doc.paragraphs)
        if para_count >= 40:
            print(f"PASS: Component 2 — Paragraph count is {para_count} (>= 40) (0.3 pts)")
            total_score += 0.3
        elif para_count > 12:
            # Partial credit: some breaks were converted
            # Scale from 0 to 0.3 based on how close to 47
            fraction = min((para_count - 12) / (40 - 12), 1.0)
            partial = 0.3 * fraction
            total_score += partial
            print(f"PARTIAL: Component 2 — Paragraph count is {para_count} "
                  f"(between 12 and 40). Partial: {partial:.2f} pts")
        else:
            print(f"FAIL: Component 2 — Paragraph count is {para_count} "
                  f"(expected >= 40, initial was 12)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Text content preserved AND line breaks removed (0.2 points)
    # This component only awards points when BOTH the core task is done (no line breaks)
    # AND the text content is intact. This ensures it doesn't score on initial_env.
    try:
        full_text = get_full_text(doc)
        text_len = len(full_text)
        # The original text is approximately 3322 chars (from both VMs).
        # Allow +/- 5% tolerance for whitespace normalization differences.
        expected_len = 3322
        tolerance = 0.05
        lower = int(expected_len * (1 - tolerance))
        upper = int(expected_len * (1 + tolerance))

        if lb_count == 0 and lower <= text_len <= upper:
            print(f"PASS: Component 3 — Line breaks removed AND text content preserved "
                  f"({text_len} chars, expected ~{expected_len}) (0.2 pts)")
            total_score += 0.2
        elif lb_count > 0:
            print(f"FAIL: Component 3 — Line breaks still present ({lb_count}), "
                  f"cannot award content preservation points")
        else:
            print(f"FAIL: Component 3 — Text content length {text_len} "
                  f"outside expected range [{lower}, {upper}]")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
