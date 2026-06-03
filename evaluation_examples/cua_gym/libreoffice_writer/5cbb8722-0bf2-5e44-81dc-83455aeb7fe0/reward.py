"""
Reward Script: Apply different first page style (no header/footer on cover page)
Task ID: writer_tech_021
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): titlePg element present — enables different first page header/footer
  Component 2 (0.3): First page header is empty (no text)
  Component 3 (0.2): First page footer is empty (no text, no field codes)
  Component 4 (0.2): Default header/footer still have content for subsequent pages
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_021'


def persist_app_state(domain):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        import time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if len(doc.sections) < 1:
        print("CRITICAL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    sec = doc.sections[0]

    # Component 1: titlePg element present in sectPr (0.3 points)
    # This enables "Different First Page" header/footer in Writer.
    # In initial_env this element is ABSENT; in golden_env it is PRESENT.
    try:
        sectPr = sec._sectPr
        titlePg = sectPr.find(qn('w:titlePg'))
        if titlePg is not None:
            print(f"PASS: Component 1 — titlePg element found in section properties (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — titlePg element NOT found in section properties")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: First page header is empty (0.3 points)
    # In initial_env, first page header is_linked_to_previous=True (inherits default header text).
    # In golden_env, first page header is_linked_to_previous=False and text is empty.
    try:
        first_hdr = sec.first_page_header
        first_hdr_linked = first_hdr.is_linked_to_previous
        first_hdr_text = "".join(p.text.strip() for p in first_hdr.paragraphs)

        if not first_hdr_linked and first_hdr_text == "":
            print(f"PASS: Component 2 — First page header is empty and not linked (0.3 pts)")
            total_score += 0.3
        elif first_hdr_linked:
            print(f"FAIL: Component 2 — First page header is still linked to previous (inherits default header)")
        else:
            print(f"FAIL: Component 2 — First page header has text: {repr(first_hdr_text)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: First page footer is empty (0.2 points)
    # In initial_env, first page footer is_linked_to_previous=True (inherits default footer with page num).
    # In golden_env, first page footer is_linked_to_previous=False and is empty (no text, no field codes).
    try:
        first_ftr = sec.first_page_footer
        first_ftr_linked = first_ftr.is_linked_to_previous
        first_ftr_text = "".join(p.text.strip() for p in first_ftr.paragraphs)

        # Also check for field codes (PAGE fields) in the footer XML
        ftr_xml = first_ftr._element.xml
        has_field_codes = 'instrText' in ftr_xml or 'fldChar' in ftr_xml

        if not first_ftr_linked and first_ftr_text == "" and not has_field_codes:
            print(f"PASS: Component 3 — First page footer is empty and not linked (0.2 pts)")
            total_score += 0.2
        elif first_ftr_linked:
            print(f"FAIL: Component 3 — First page footer is still linked to previous (inherits default footer)")
        elif first_ftr_text != "" or has_field_codes:
            print(f"FAIL: Component 3 — First page footer has content: text={repr(first_ftr_text)}, field_codes={has_field_codes}")
        else:
            print(f"FAIL: Component 3 — First page footer check failed")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Default header and footer still have content (0.2 points)
    # The task says "the rest of the document retains" header/footer.
    # Default header should still have document title text.
    # Default footer should still have page number content.
    # This component ONLY scores if titlePg is enabled (Component 1 passed),
    # ensuring we measure the task-introduced state, not a precondition.
    try:
        if titlePg is not None:
            default_hdr = sec.header
            default_hdr_text = "".join(p.text.strip() for p in default_hdr.paragraphs)

            default_ftr = sec.footer
            default_ftr_text = "".join(p.text.strip() for p in default_ftr.paragraphs)
            # Check for PAGE field codes in default footer XML
            ftr_xml_default = default_ftr._element.xml
            has_page_field = 'PAGE' in ftr_xml_default or 'instrText' in ftr_xml_default

            header_ok = len(default_hdr_text) > 0
            footer_ok = len(default_ftr_text) > 0 or has_page_field

            if header_ok and footer_ok:
                print(f"PASS: Component 4 — Default header has text ({repr(default_hdr_text[:50])}) and footer has content (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 4 — Default header text={repr(default_hdr_text)}, footer text={repr(default_ftr_text)}, has_page_field={has_page_field}")
        else:
            print(f"FAIL: Component 4 — Skipped because titlePg is not enabled (Component 1 failed)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint: persist app state then verify
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
