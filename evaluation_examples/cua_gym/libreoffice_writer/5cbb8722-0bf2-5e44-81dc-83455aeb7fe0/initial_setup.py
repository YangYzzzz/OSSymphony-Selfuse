"""
Initial Setup: Create a Writer document with cover page and content pages,
with headers and footers enabled on ALL pages (no first-page differentiation).
Task ID: writer_tech_021
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_021'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Ensure NO "different first page" is set ---
    # Remove titlePg if it exists (ensures all pages share same header/footer)
    sectPr = section._sectPr
    for titlePg in sectPr.findall(qn('w:titlePg')):
        sectPr.remove(titlePg)

    # --- Header (same on ALL pages, including first) ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    run = hp.add_run("Technical Architecture Report")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Footer with page number (same on ALL pages) ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run_prefix = fp.add_run("Page ")
    run_prefix.font.size = Pt(9)
    run_prefix.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # PAGE field code
    r1 = fp.add_run()
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    fldChar1 = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fldChar1)

    r2 = fp.add_run()
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    instrText = r2._element.makeelement(qn('w:instrText'), {})
    instrText.text = ' PAGE '
    r2._element.append(instrText)

    r3 = fp.add_run()
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    fldChar2 = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fldChar2)

    # === COVER PAGE ===
    # Add some vertical space
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_heading("Technical Architecture Report", level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_sub = subtitle.add_run("Cloud Infrastructure Modernization Project")
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph("")

    info_para = doc.add_paragraph()
    info_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_info = info_para.add_run("Prepared by: Elena Vasquez, Lead Architect\nRevision 3.2 — March 2026\nConfidential — Internal Use Only")
    run_info.font.size = Pt(11)
    run_info.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Page break after cover
    doc.add_page_break()

    # === PAGE 2: Executive Summary ===
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This document outlines the proposed technical architecture for the cloud "
        "infrastructure modernization initiative at Meridian Technologies. The project "
        "aims to migrate legacy on-premise systems to a hybrid cloud environment, "
        "leveraging containerized microservices and managed database solutions."
    )
    doc.add_paragraph(
        "The migration will proceed in three phases over an 18-month timeline. Phase 1 "
        "focuses on identity and access management, Phase 2 addresses compute and storage "
        "migration, and Phase 3 completes the transition with networking and monitoring "
        "infrastructure."
    )

    # === System Overview ===
    doc.add_heading("2. System Overview", level=1)
    doc.add_paragraph(
        "The current infrastructure consists of 47 physical servers across two data centers "
        "in Portland and Chicago. These servers run a mix of Red Hat Enterprise Linux 8, "
        "Windows Server 2019, and several legacy Solaris 11 instances that support the "
        "core ERP platform."
    )

    doc.add_heading("2.1 Current Architecture", level=2)
    doc.add_paragraph(
        "The existing architecture follows a three-tier model with dedicated web servers, "
        "application servers, and database servers. Inter-tier communication occurs over "
        "a private 10GbE backbone with hardware load balancers managing traffic distribution."
    )

    # Add a table for server inventory
    doc.add_paragraph("")
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["Server Role", "Count", "OS", "Avg. Load (%)"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    data = [
        ["Web Servers", "8", "RHEL 8", "62"],
        ["App Servers", "14", "RHEL 8 / Win 2019", "78"],
        ["Database Servers", "6", "RHEL 8", "85"],
        ["Legacy ERP", "4", "Solaris 11", "71"],
        ["Utility / CI-CD", "15", "RHEL 8", "45"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_heading("2.2 Target Architecture", level=2)
    doc.add_paragraph(
        "The target architecture adopts a Kubernetes-based microservices platform running "
        "on AWS EKS, with Aurora PostgreSQL replacing the existing Oracle RAC cluster. "
        "A service mesh (Istio) will handle inter-service communication, mTLS, and "
        "observability."
    )

    # === Migration Strategy ===
    doc.add_heading("3. Migration Strategy", level=1)
    doc.add_paragraph(
        "The migration follows the 6R framework: Rehost, Replatform, Refactor, Repurchase, "
        "Retain, and Retire. Each workload has been assessed and categorized into one of "
        "these strategies based on business criticality, technical debt, and cloud readiness."
    )

    doc.add_heading("3.1 Phase 1 — Identity & Access (Months 1–4)", level=2)
    doc.add_paragraph(
        "Migrate Active Directory to AWS Directory Service. Implement SSO via Okta with "
        "SAML 2.0 federation. Establish VPN tunnels between on-premise and AWS VPCs. "
        "Deploy HashiCorp Vault for secrets management."
    )

    doc.add_heading("3.2 Phase 2 — Compute & Storage (Months 5–12)", level=2)
    doc.add_paragraph(
        "Containerize stateless application services using Docker. Deploy Kubernetes "
        "clusters with auto-scaling node groups. Migrate file storage to S3 with "
        "cross-region replication. Convert Oracle databases to Aurora PostgreSQL using "
        "AWS DMS with ongoing replication."
    )

    doc.add_heading("3.3 Phase 3 — Networking & Monitoring (Months 13–18)", level=2)
    doc.add_paragraph(
        "Implement AWS Transit Gateway for centralized routing. Deploy CloudFront CDN "
        "for global content delivery. Set up Prometheus/Grafana monitoring stack with "
        "PagerDuty integration. Conduct final cutover and decommission legacy hardware."
    )

    # === Risk Assessment ===
    doc.add_heading("4. Risk Assessment", level=1)
    doc.add_paragraph(
        "Key risks include data loss during Oracle-to-PostgreSQL migration, extended "
        "downtime during DNS cutover, and compatibility issues with the Solaris-based "
        "ERP modules. Mitigation strategies involve parallel-run periods, automated "
        "rollback procedures, and dedicated compatibility testing environments."
    )

    doc.add_paragraph(
        "The project budget is estimated at $2.4 million with a contingency reserve of "
        "15%. Expected annual savings post-migration are $680,000 from reduced hardware "
        "maintenance, licensing, and data center co-location fees."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
