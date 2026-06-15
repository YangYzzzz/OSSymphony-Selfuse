"""
Initial Setup: PTO Tracking Sheet 2026 - Employee info header only (no table)
Task ID: writer_hr_042
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_042'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('PTO Tracking Sheet 2026', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Paid Time Off Accrual & Usage Record')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)
    run.italic = True

    doc.add_paragraph()  # spacer

    # --- Employee Information Section ---
    info_heading = doc.add_heading('Employee Information', level=2)

    # Employee info fields (to be filled in)
    fields = [
        ('Employee Name', 'Rachel Nguyen'),
        ('Employee ID', 'EMP-2024-0387'),
        ('Department', 'Product Development'),
        ('Manager', 'David Kowalski'),
        ('Hire Date', 'March 12, 2024'),
        ('Annual PTO Entitlement', '15 days'),
        ('Monthly Accrual Rate', '1.25 days'),
    ]

    for label, value in fields:
        para = doc.add_paragraph()
        label_run = para.add_run(f'{label}: ')
        label_run.bold = True
        label_run.font.size = Pt(11)
        value_run = para.add_run(value)
        value_run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # --- Instructions paragraph ---
    instructions = doc.add_heading('Instructions', level=2)

    inst_para = doc.add_paragraph()
    inst_run = inst_para.add_run(
        'Please create a PTO tracking table below this section to track monthly '
        'paid time off accrual and usage for the year 2026. The table should include '
        'columns for Month, Days Available, Days Used, and Days Remaining, with a '
        'total row at the bottom.'
    )
    inst_run.font.size = Pt(11)

    doc.add_paragraph()  # spacer for where the table should go

    # --- Footer note ---
    note_para = doc.add_paragraph()
    note_run = note_para.add_run(
        'Note: PTO requests must be submitted at least 5 business days in advance '
        'and approved by your direct manager. Unused PTO days may be carried over '
        'to the next calendar year, up to a maximum of 5 days.'
    )
    note_run.font.size = Pt(9)
    note_run.italic = True
    note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
