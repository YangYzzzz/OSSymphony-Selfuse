"""
Initial Setup: Statistics methodology document with p-hacking paragraph
Task ID: writer_bs_044
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_044'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('Methodological Rigor in Modern Statistical Research', level=1)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=2)

    p1 = doc.add_paragraph()
    p1.add_run(
        'The replication crisis in psychology and biomedical sciences has prompted '
        'a fundamental re-examination of how statistical analyses are conducted and '
        'reported. Over the past decade, researchers have identified numerous practices '
        'that inflate false-positive rates, undermine the credibility of published findings, '
        'and erode public trust in scientific institutions.'
    )

    p2 = doc.add_paragraph()
    p2.add_run(
        'This paper reviews the most prominent threats to statistical validity, including '
        'selective reporting, undisclosed flexibility in data collection, and the broader '
        'cultural incentives that perpetuate questionable research practices (QRPs). We '
        'propose a framework for evaluating methodological transparency that builds on '
        'existing open-science initiatives.'
    )

    # --- Section 2: The Problem of P-Hacking ---
    doc.add_heading('2. The Problem of P-Hacking', level=2)

    # THE KEY PARAGRAPH - contains the text that needs footnote/endnote
    p3 = doc.add_paragraph()
    p3.add_run(
        'Among the most pervasive concerns in contemporary research methodology is '
        'the practice of p-hacking has been widely criticized in the literature '
        '(Simmons et al., 2011). This involves the selective analysis of data\u2014trying '
        'multiple statistical tests, excluding outliers post hoc, or adjusting sample '
        'sizes\u2014until a desired level of statistical significance is achieved. The '
        'consequences of such practices extend far beyond individual studies, contributing '
        'to a systematic overestimation of effect sizes across entire research domains.'
    )

    p4 = doc.add_paragraph()
    p4.add_run(
        'Several large-scale replication projects, including the Reproducibility Project: '
        'Psychology (Open Science Collaboration, 2015) and the Many Labs initiative '
        '(Klein et al., 2018), have demonstrated that a substantial proportion of published '
        'findings fail to replicate under pre-registered conditions. While not all replication '
        'failures can be attributed to p-hacking, the practice remains a central concern in '
        'discussions of research integrity.'
    )

    # --- Section 3: Mitigation Strategies ---
    doc.add_heading('3. Mitigation Strategies', level=2)

    p5 = doc.add_paragraph()
    p5.add_run(
        'Pre-registration of hypotheses and analysis plans has emerged as one of the '
        'most promising solutions to the p-hacking problem. By committing to a specific '
        'analytical approach before data collection begins, researchers limit their degrees '
        'of freedom and make it more difficult to engage in post hoc rationalization of '
        'unexpected findings.'
    )

    p6 = doc.add_paragraph()
    p6.add_run(
        'Additionally, journals have begun adopting Registered Reports, a publication '
        'format in which manuscripts are peer-reviewed and accepted based on the quality '
        'of the research question and methodology, prior to data collection. This approach '
        'decouples publication decisions from the statistical significance of results, '
        'reducing the incentive to manipulate analyses.'
    )

    # --- Section 4: Conclusion ---
    doc.add_heading('4. Conclusion', level=2)

    p7 = doc.add_paragraph()
    p7.add_run(
        'Addressing the challenges posed by p-hacking and related questionable research '
        'practices requires a multi-faceted approach involving institutional reform, '
        'methodological education, and technological tools for enhancing transparency. '
        'The scientific community has made meaningful progress in recent years, but '
        'sustained effort is needed to establish a culture of openness and accountability '
        'in empirical research.'
    )

    # --- References section ---
    doc.add_heading('References', level=2)

    refs = [
        'Klein, R. A., Vianello, M., Hasselman, F., et al. (2018). Many Labs 2: Investigating variation in replicability across samples and settings. Advances in Methods and Practices in Psychological Science, 1(4), 443-490.',
        'Open Science Collaboration. (2015). Estimating the reproducibility of psychological science. Science, 349(6251), aac4716.',
        'Simmons, J. P., Nelson, L. D., & Simonsohn, U. (2011). False-Positive Psychology: Undisclosed Flexibility in Data Collection and Analysis Allows Presenting Anything as Significant. Psychological Science, 22(11), 1359-1366.',
    ]
    for ref in refs:
        rp = doc.add_paragraph()
        rp.add_run(ref)
        rp.paragraph_format.left_indent = Inches(0.5)
        rp.paragraph_format.first_line_indent = Inches(-0.5)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
