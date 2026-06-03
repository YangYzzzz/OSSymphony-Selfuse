"""
Initial Setup: Split first paragraph into sentences with spacing
Task ID: wrpara_015
Domain: libreoffice_writer

Creates a 3-paragraph document about renewable energy.
Paragraph 1 has 5 sentences in one block.
Paragraphs 2 and 3 are separate topics.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'wrpara_015'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Paragraph 1: 5 sentences about renewable energy in ONE block ---
    p1_text = (
        "Renewable energy sources have become increasingly important in the global effort to reduce carbon emissions and combat climate change. "
        "Solar power installations have grown by over 40 percent annually in the past decade, making it one of the fastest-growing energy sectors worldwide. "
        "Wind energy now accounts for approximately 8 percent of total electricity generation in the United States and continues to expand across rural and offshore regions. "
        "Governments around the world have implemented subsidies and tax incentives to encourage businesses and homeowners to adopt clean energy technologies. "
        "The transition to renewable energy is expected to create millions of new jobs in manufacturing, installation, and maintenance over the next twenty years."
    )
    para1 = doc.add_paragraph(p1_text)

    # --- Paragraph 2: Energy storage challenges ---
    p2_text = (
        "Energy storage remains one of the most significant challenges facing the renewable energy industry. "
        "Battery technology has improved dramatically, with lithium-ion costs dropping by nearly 90 percent since 2010, "
        "but large-scale grid storage solutions are still in development. Pumped hydroelectric storage currently "
        "accounts for over 95 percent of global energy storage capacity, though newer technologies like compressed "
        "air and hydrogen fuel cells are gaining traction. Researchers at institutions such as MIT and Stanford "
        "are exploring solid-state batteries that could revolutionize how we store and distribute clean energy."
    )
    para2 = doc.add_paragraph(p2_text)

    # --- Paragraph 3: Economic impact ---
    p3_text = (
        "The economic impact of the renewable energy transition extends far beyond the energy sector itself. "
        "Communities that were once dependent on coal mining have begun retraining workers for careers in solar "
        "panel manufacturing and wind turbine maintenance. Investment in clean energy infrastructure reached a "
        "record $495 billion globally in 2024, with China, the European Union, and the United States leading "
        "the way. Small businesses are also benefiting from reduced electricity costs after installing rooftop "
        "solar systems, with average savings of 30 to 50 percent on monthly utility bills."
    )
    para3 = doc.add_paragraph(p3_text)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
