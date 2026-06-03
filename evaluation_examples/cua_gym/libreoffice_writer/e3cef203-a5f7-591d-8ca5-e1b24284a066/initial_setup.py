"""
Initial Setup: Sourdough Blog Post - Unformatted Draft
Task ID: writer_creative_022
Domain: libreoffice_writer

Creates a blog post draft about sourdough bread making with all content
in default paragraph style (12pt), no heading styles, and no bold formatting.
The agent will need to apply heading styles, and bold ingredient measurements.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'  # Task context says ~/Desktop/
TASK_ID = 'sourdough_blog'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure the Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Helper to add a plain paragraph with 12pt, Default Paragraph Style
    def add_plain(text, size_pt=12):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(size_pt)
        run.bold = False
        run.italic = False
        return para

    # ----- Title (plain text, no Heading style) -----
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("The Art of Sourdough: A Beginner's Journey")
    title_run.font.size = Pt(12)
    title_run.bold = False

    # ----- Section: Why Sourdough? -----
    add_plain("Why Sourdough?")

    add_plain(
        "Sourdough bread is one of the oldest and most rewarding baking traditions in the world. "
        "Unlike commercial yeast breads, sourdough relies on wild yeast and bacteria naturally present "
        "in flour and the environment. The result is a complex, tangy flavor that no store-bought bread "
        "can replicate."
    )
    add_plain(
        "Beyond the taste, sourdough offers real nutritional benefits. The long fermentation process "
        "breaks down phytic acid and partially predigests gluten, making it easier for many people to "
        "digest. Once you bake your first successful loaf, you'll understand why so many bakers call "
        "sourdough a lifelong obsession."
    )

    # ----- Section: Getting Started: Your First Starter -----
    add_plain("Getting Started: Your First Starter")

    add_plain(
        "A sourdough starter is a live culture of wild yeast and lactic acid bacteria. You can create "
        "one from scratch using nothing more than flour and water — it typically takes 5 to 7 days for "
        "the culture to become active and reliable. Mix 50g of whole wheat or rye flour with 50ml of "
        "lukewarm water in a clean jar, cover loosely, and leave at room temperature."
    )
    add_plain(
        "Every 24 hours, discard half the starter and feed it with fresh flour and water. Within a few "
        "days you should see bubbles forming, a sign that the wild yeast is thriving. Your starter is "
        "ready to bake with when it reliably doubles in size within 4 to 8 hours of feeding and has a "
        "pleasant, mildly sour aroma."
    )

    # ----- Section: The Basic Recipe -----
    add_plain("The Basic Recipe")

    add_plain(
        "This recipe makes one medium-sized country loaf. Gather your ingredients before you begin:"
    )

    # Ingredient lines — measurements are NOT bold in initial state
    ingredients = [
        "500g bread flour",
        "350ml warm water",
        "100g active starter",
        "10g sea salt",
    ]
    for ingredient in ingredients:
        para = doc.add_paragraph()
        run = para.add_run(ingredient)
        run.font.size = Pt(12)
        run.bold = False

    # ----- Section: The Baking Process -----
    add_plain("The Baking Process")

    add_plain(
        "Step 1 — Autolyse: Combine the bread flour and 300ml of the warm water. Mix until no dry "
        "flour remains, then cover and rest for 30 to 60 minutes. This rest period, called autolyse, "
        "allows the flour to fully hydrate and kick-starts gluten development without any kneading."
    )
    add_plain(
        "Step 2 — Add Starter and Salt: Add the active starter and dissolve the sea salt in the "
        "remaining 50ml of warm water, then pour everything into the dough. Use your hand to fold "
        "and squeeze the dough until the starter and salt are fully incorporated, about 3 to 5 minutes."
    )
    add_plain(
        "Step 3 — Bulk Fermentation: Cover the dough and let it ferment at room temperature for "
        "4 to 6 hours. During this time, perform a set of stretch-and-folds every 30 minutes for "
        "the first 2 hours. The dough should increase in volume by roughly 50% and feel light and "
        "airy when bulk fermentation is complete."
    )
    add_plain(
        "Step 4 — Shape, Proof, and Bake: Shape the dough into a tight round or oval and place it "
        "seam-side up in a floured banneton or bowl lined with a floured cloth. Refrigerate overnight "
        "(8 to 16 hours) for a cold proof. When ready to bake, preheat your oven to 250°C with a "
        "Dutch oven inside. Bake covered for 20 minutes, then uncovered for 20 to 25 minutes until "
        "deep golden brown."
    )

    # ----- Section: Tips for Success -----
    add_plain("Tips for Success")

    add_plain(
        "Use a kitchen scale for all measurements. Sourdough is a precision craft and volume "
        "measurements of flour are notoriously inconsistent. Even a 10% variation in hydration "
        "can dramatically change the texture of your final loaf."
    )
    add_plain(
        "Temperature matters more than time. Fermentation speed is directly controlled by ambient "
        "temperature. A warm kitchen (24–26°C) can cut bulk fermentation to 4 hours, while a cool "
        "kitchen (18–20°C) may need 8 hours. Learn to read the dough, not the clock."
    )
    add_plain(
        "Keep your starter consistent. Feed your starter at the same time each day, using the same "
        "flour and water ratio. Consistency trains the microbial community to work predictably, "
        "giving you reliable results every bake."
    )
    add_plain(
        "Score confidently and decisively. Before loading the dough into the oven, make a swift "
        "slash with a sharp lame or razor blade. Hesitant scoring tears the surface; one clean, "
        "angled cut about 1cm deep will allow beautiful oven spring."
    )
    add_plain(
        "Embrace failure as a learning tool. Your first few loaves may be dense, gummy, or under-"
        "proofed — and that is completely normal. Take notes after every bake, adjust one variable "
        "at a time, and within 5 to 10 attempts you will be producing loaves you are proud to share."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
