"""
Initial Setup: Mailing list with duplicate entries for duplicate removal task
Task ID: osworld_writer_duplicate_line_removal_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_duplicate_line_removal_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('Community Newsletter Mailing List')
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Subtitle/intro
    intro = doc.add_paragraph()
    intro_run = intro.add_run('Subscriber Addresses — Spring 2025 Edition')
    intro_run.italic = True
    intro_run.font.size = Pt(11)
    intro.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # blank line

    # 37 unique addresses (first occurrences)
    unique_addresses = [
        '742 Evergreen Terrace, Springfield, IL 62704',
        '1600 Pennsylvania Ave NW, Washington, DC 20500',
        '221B Baker Street, London, NW1 6XE, UK',
        '350 Fifth Avenue, New York, NY 10118',
        '1 Infinite Loop, Cupertino, CA 95014',
        '4059 Mt Lee Dr, Hollywood, CA 90068',
        '800 N Michigan Ave, Chicago, IL 60611',
        '1 Harbor Drive, San Diego, CA 92101',
        '500 W 2nd St, Austin, TX 78701',
        '55 Water Street, New York, NY 10041',
        '1234 Oak Hill Blvd, Nashville, TN 37201',
        '88 Pine Street, Portland, OR 97201',
        '310 Maple Avenue, Denver, CO 80203',
        '9900 Wilshire Blvd, Beverly Hills, CA 90210',
        '47 Riverside Road, Boston, MA 02101',
        '623 Sunset Strip, Los Angeles, CA 90028',
        '1 Embarcadero Center, San Francisco, CA 94111',
        '3400 N Lake Shore Dr, Chicago, IL 60657',
        '890 Peachtree St NE, Atlanta, GA 30309',
        '2200 Mission College Blvd, Santa Clara, CA 95054',
        '150 Broadway, New York, NY 10038',
        '7 World Trade Center, New York, NY 10007',
        '1000 Main Street, Houston, TX 77002',
        '412 Elm Street, Cincinnati, OH 45202',
        '675 Ponce De Leon Ave NE, Atlanta, GA 30308',
        '330 N Wabash Ave, Chicago, IL 60611',
        '1 Microsoft Way, Redmond, WA 98052',
        '1600 Amphitheatre Pkwy, Mountain View, CA 94043',
        '2100 Geng Road, Palo Alto, CA 94303',
        '1601 S California Ave, Palo Alto, CA 94304',
        '500 Oracle Pkwy, Redwood Shores, CA 94065',
        '901 Cherry Ave, San Bruno, CA 94066',
        '770 Broadway, New York, NY 10003',
        '1455 Market St, San Francisco, CA 94103',
        '410 Terry Ave N, Seattle, WA 98109',
        '2380 McGaw Ave, Irvine, CA 92614',
        '100 Universal City Plaza, Universal City, CA 91608',
    ]

    # 15 duplicate entries to scatter throughout the list
    # These are exact copies of specific addresses above
    duplicates = [
        (4, '1 Infinite Loop, Cupertino, CA 95014'),           # dup of index 4
        (9, '55 Water Street, New York, NY 10041'),             # dup of index 9
        (0, '742 Evergreen Terrace, Springfield, IL 62704'),    # dup of index 0
        (14, '47 Riverside Road, Boston, MA 02101'),            # dup of index 14
        (19, '2200 Mission College Blvd, Santa Clara, CA 95054'),  # dup of index 19
        (2, '221B Baker Street, London, NW1 6XE, UK'),          # dup of index 2
        (26, '1 Microsoft Way, Redmond, WA 98052'),             # dup of index 26
        (11, '88 Pine Street, Portland, OR 97201'),             # dup of index 11
        (33, '1455 Market St, San Francisco, CA 94103'),        # dup of index 33
        (6, '800 N Michigan Ave, Chicago, IL 60611'),           # dup of index 6
        (22, '1000 Main Street, Houston, TX 77002'),            # dup of index 22
        (27, '1600 Amphitheatre Pkwy, Mountain View, CA 94043'),  # dup of index 27
        (16, '1 Embarcadero Center, San Francisco, CA 94111'),  # dup of index 16
        (30, '500 Oracle Pkwy, Redwood Shores, CA 94065'),      # dup of index 30
        (35, '2380 McGaw Ave, Irvine, CA 92614'),               # dup of index 35
    ]

    # Build the full 52-entry list by inserting duplicates at specific positions
    # We'll place duplicates after certain unique entries to scatter them
    # Insert positions in the final list (after building unique list):
    # Position mapping: insert after the Nth entry in the growing list
    # Strategy: interleave duplicates at regular intervals across the list
    full_list = list(unique_addresses)  # start with 37 unique

    # Insert duplicates at specific positions to create a 52-entry list
    # Insert positions (0-indexed into the list at time of insertion):
    insert_positions = [
        (5, '1 Infinite Loop, Cupertino, CA 95014'),
        (10, '55 Water Street, New York, NY 10041'),
        (15, '742 Evergreen Terrace, Springfield, IL 62704'),
        (18, '47 Riverside Road, Boston, MA 02101'),
        (21, '2200 Mission College Blvd, Santa Clara, CA 95054'),
        (24, '221B Baker Street, London, NW1 6XE, UK'),
        (27, '1 Microsoft Way, Redmond, WA 98052'),
        (29, '88 Pine Street, Portland, OR 97201'),
        (31, '1455 Market St, San Francisco, CA 94103'),
        (33, '800 N Michigan Ave, Chicago, IL 60611'),
        (35, '1000 Main Street, Houston, TX 77002'),
        (38, '1600 Amphitheatre Pkwy, Mountain View, CA 94043'),
        (41, '1 Embarcadero Center, San Francisco, CA 94111'),
        (44, '500 Oracle Pkwy, Redwood Shores, CA 94065'),
        (48, '2380 McGaw Ave, Irvine, CA 92614'),
    ]

    for pos, addr in insert_positions:
        full_list.insert(pos, addr)

    # Verify we have 52 entries
    assert len(full_list) == 52, f'Expected 52 entries, got {len(full_list)}'

    # Write each address as a paragraph
    for addr in full_list:
        p = doc.add_paragraph()
        run = p.add_run(addr)
        run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')
    print(f'Total entries: {len(full_list)}')

    # Verify unique count
    seen = set()
    unique_count = 0
    for addr in full_list:
        if addr not in seen:
            seen.add(addr)
            unique_count += 1
    print(f'Unique entries: {unique_count}')
    print(f'Duplicate entries: {len(full_list) - unique_count}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
