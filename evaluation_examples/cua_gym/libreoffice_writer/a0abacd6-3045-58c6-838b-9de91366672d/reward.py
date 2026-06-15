"""
Reward Script: Set a right-aligned tab stop at 16 cm with dot leaders for the table of contents entries.
Task ID: writer_para_011
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): All 6 TOC entry paragraphs (para index 1-6) have a right-aligned tab stop at ~16 cm
  Component 2 (0.4): All paragraphs 1-6 have DOTS leader on that right-aligned tab stop
"""

import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

WORKDIR = '/home/user'
TASK_ID = 'writer_para_011'

# 16 cm in EMU (English Metric Units) — python-docx stores positions in EMU
EXPECTED_POSITION_EMU = Cm(16)     # 5760000 EMU
POSITION_TOLERANCE = 5000          # ~0.5mm tolerance for rounding differences

# Indices of the TOC entry paragraphs (Introduction, Project Background, ..., Appendices)
TOC_ENTRY_INDICES = [1, 2, 3, 4, 5, 6]
# Expected text snippets to confirm we're looking at the right paragraphs
EXPECTED_TEXTS = [
    'Introduction',
    'Project Background',
    'Methodology',
    'Results and Analysis',
    'Conclusions',
    'Appendices',
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition gate: load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: check document has at least 7 paragraphs
    if len(doc.paragraphs) < 7:
        print(f"CRITICAL: Expected at least 7 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: verify text content is intact (no corruption)
    content_ok = True
    for idx, expected_text in zip(TOC_ENTRY_INDICES, EXPECTED_TEXTS):
        para_text = doc.paragraphs[idx].text
        if expected_text not in para_text:
            print(f"CRITICAL: Para {idx} text mismatch. Expected to contain '{expected_text}', got '{para_text}'")
            content_ok = False
    if not content_ok:
        print("CRITICAL: Text content corrupted — aborting scoring")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All 6 TOC entry paragraphs have a right-aligned tab stop at ~16 cm (0.6 points)
    # This FAILS on initial_env (no tab stops) and PASSES on golden_env (tab stops added)
    try:
        paras_with_right_tab = 0
        for idx in TOC_ENTRY_INDICES:
            para = doc.paragraphs[idx]
            tab_stops = list(para.paragraph_format.tab_stops)
            has_right_tab_16cm = False
            for ts in tab_stops:
                # Skip default/clear stops
                if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
                    continue
                if ts.alignment == WD_TAB_ALIGNMENT.LEFT and ts.position == 0:
                    continue
                # Check for right-aligned tab near 16 cm
                if (ts.alignment == WD_TAB_ALIGNMENT.RIGHT and
                        abs(ts.position - EXPECTED_POSITION_EMU) <= POSITION_TOLERANCE):
                    has_right_tab_16cm = True
                    break
            if has_right_tab_16cm:
                paras_with_right_tab += 1

        if paras_with_right_tab == len(TOC_ENTRY_INDICES):
            print(f"PASS: Component 1 — All {len(TOC_ENTRY_INDICES)} TOC entry paragraphs have right-aligned tab at ~16 cm (0.6 pts)")
            total_score += 0.6
        elif paras_with_right_tab > 0:
            partial = round(0.6 * paras_with_right_tab / len(TOC_ENTRY_INDICES), 4)
            print(f"PARTIAL: Component 1 — {paras_with_right_tab}/{len(TOC_ENTRY_INDICES)} paragraphs have right-aligned tab at ~16 cm ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No TOC entry paragraph has a right-aligned tab stop at ~16 cm")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All 6 TOC entry paragraphs have DOTS leader on that right-aligned tab stop (0.4 points)
    # This FAILS on initial_env (no tab stops) and PASSES on golden_env (tab stops with DOTS added)
    try:
        paras_with_dots_leader = 0
        for idx in TOC_ENTRY_INDICES:
            para = doc.paragraphs[idx]
            tab_stops = list(para.paragraph_format.tab_stops)
            has_dots_on_right_tab = False
            for ts in tab_stops:
                if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
                    continue
                if ts.alignment == WD_TAB_ALIGNMENT.LEFT and ts.position == 0:
                    continue
                # Check for right-aligned tab near 16 cm WITH dots leader
                if (ts.alignment == WD_TAB_ALIGNMENT.RIGHT and
                        abs(ts.position - EXPECTED_POSITION_EMU) <= POSITION_TOLERANCE and
                        ts.leader == WD_TAB_LEADER.DOTS):
                    has_dots_on_right_tab = True
                    break
            if has_dots_on_right_tab:
                paras_with_dots_leader += 1

        if paras_with_dots_leader == len(TOC_ENTRY_INDICES):
            print(f"PASS: Component 2 — All {len(TOC_ENTRY_INDICES)} TOC entry paragraphs have DOTS leader on right-aligned tab (0.4 pts)")
            total_score += 0.4
        elif paras_with_dots_leader > 0:
            partial = round(0.4 * paras_with_dots_leader / len(TOC_ENTRY_INDICES), 4)
            print(f"PARTIAL: Component 2 — {paras_with_dots_leader}/{len(TOC_ENTRY_INDICES)} paragraphs have DOTS leader on right-aligned tab ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No TOC entry paragraph has a DOTS leader on the right-aligned tab stop")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM environment
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
