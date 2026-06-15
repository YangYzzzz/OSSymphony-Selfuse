"""
Initial Setup: Set a right-aligned tab stop at 16 cm with dot leaders for table of contents entries.
Task ID: writer_para_011
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_para_011'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: 'Table of Contents' — Heading 1, center-aligned
    heading = doc.add_heading('Table of Contents', level=1)
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Paragraphs 2–7: TOC entries with tab separator (no tab stops set)
    toc_entries = [
        'Introduction\t5',
        'Project Background\t8',
        'Methodology\t15',
        'Results and Analysis\t22',
        'Conclusions\t35',
        'Appendices\t40',
    ]
    for entry in toc_entries:
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # No tab stops added — that is the task for the agent
        run = para.add_run(entry)
        run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer with DISPLAY=:0
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
