"""
Initial Setup: Event flyer document with banner image as separate file
Task ID: writer_obj_051
Domain: libreoffice_writer

Creates:
  - /home/user/Desktop/event_flyer.docx  (single-page event announcement, no banner inserted)
  - /home/user/Desktop/banner.jpg         (1920x400 pixel banner image)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_051'
DOC_OUTPUT = f'{WORKDIR}/event_flyer.docx'
IMG_OUTPUT = f'{WORKDIR}/banner.jpg'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_banner_image():
    """Create a 1920x400 pixel banner image."""
    img = Image.new('RGB', (1920, 400), color=(41, 98, 175))
    draw = ImageDraw.Draw(img)

    # Draw a gradient-like background with rectangles
    for i in range(0, 1920, 10):
        shade = int(41 + (i / 1920) * 60)
        draw.rectangle([i, 0, i + 10, 400], fill=(shade, 98 + int(i / 40), 175))

    # Draw decorative elements
    draw.rectangle([0, 320, 1920, 400], fill=(255, 165, 0))
    draw.rectangle([0, 0, 1920, 60], fill=(20, 50, 120))

    # Draw text if possible
    try:
        font_large = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf', 80)
        font_small = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 40)
    except Exception:
        font_large = ImageFont.load_default()
        font_small = font_large

    draw.text((960, 160), 'ANNUAL TECH SUMMIT 2025', fill=(255, 255, 255),
              font=font_large, anchor='mm')
    draw.text((960, 260), 'Innovation • Collaboration • Excellence', fill=(255, 220, 100),
              font=font_small, anchor='mm')

    img.save(IMG_OUTPUT, 'JPEG', quality=95)
    print(f'Banner image created: {IMG_OUTPUT}')


def create_initial_document():
    """Create event_flyer.docx with event content but NO banner image inserted."""
    doc = Document()

    # Set page margins (A4 with standard margins)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)

    # Event title
    title = doc.add_heading('Annual Tech Summit 2025', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x44, 0x91)
        run.font.size = Pt(28)

    # Subtitle
    subtitle = doc.add_paragraph('Empowering Innovation Through Technology')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # Event details section
    details_heading = doc.add_heading('Event Details', level=2)
    for run in details_heading.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x44, 0x91)

    details_para = doc.add_paragraph()
    details_para.add_run('Date: ').bold = True
    details_para.add_run('Saturday, June 14, 2025')

    venue_para = doc.add_paragraph()
    venue_para.add_run('Venue: ').bold = True
    venue_para.add_run('Grand Convention Center, 500 Technology Boulevard, San Francisco, CA 94105')

    time_para = doc.add_paragraph()
    time_para.add_run('Time: ').bold = True
    time_para.add_run('9:00 AM – 6:00 PM (Registration opens at 8:30 AM)')

    doc.add_paragraph()

    # About section
    about_heading = doc.add_heading('About the Event', level=2)
    for run in about_heading.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x44, 0x91)

    about_text = doc.add_paragraph(
        'Join us for the most anticipated technology event of the year! The Annual Tech Summit 2025 '
        'brings together industry leaders, innovators, and entrepreneurs for a full day of inspiring '
        'keynote speeches, hands-on workshops, and invaluable networking opportunities.'
    )
    about_text.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    # Keynote speakers
    speakers_heading = doc.add_heading('Keynote Speakers', level=2)
    for run in speakers_heading.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x44, 0x91)

    speakers = [
        ('Dr. Alexandra Rivera', 'CEO, FutureTech Innovations', 'The Future of AI in Enterprise'),
        ('Marcus Chen', 'CTO, CloudScale Systems', 'Building Resilient Distributed Architectures'),
        ('Priya Patel', 'VP Engineering, DataStream Corp', 'Real-time Analytics at Scale'),
        ('James O\'Brien', 'Founder, GreenCode Initiative', 'Sustainable Software Development'),
    ]

    for name, title_role, topic in speakers:
        speaker_para = doc.add_paragraph(style='List Bullet')
        run_name = speaker_para.add_run(f'{name}')
        run_name.bold = True
        run_name.font.color.rgb = RGBColor(0x1D, 0x44, 0x91)
        speaker_para.add_run(f', {title_role}')
        speaker_para.add_run(f'\n  Topic: {topic}').italic = True

    doc.add_paragraph()

    # Schedule section
    schedule_heading = doc.add_heading('Schedule Highlights', level=2)
    for run in schedule_heading.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x44, 0x91)

    schedule_items = [
        ('8:30 AM', 'Registration & Welcome Coffee'),
        ('9:00 AM', 'Opening Keynote: The Future of AI in Enterprise'),
        ('10:30 AM', 'Panel Discussion: Emerging Technologies in 2025'),
        ('12:00 PM', 'Networking Lunch'),
        ('1:30 PM', 'Workshop Sessions (3 parallel tracks)'),
        ('3:30 PM', 'Startup Showcase & Demo Hour'),
        ('5:00 PM', 'Closing Keynote & Awards Ceremony'),
        ('6:00 PM', 'Evening Reception'),
    ]

    for time_slot, event_name in schedule_items:
        sched_para = doc.add_paragraph(style='List Bullet')
        sched_para.add_run(f'{time_slot}  ').bold = True
        sched_para.add_run(event_name)

    doc.add_paragraph()

    # Registration section
    reg_heading = doc.add_heading('Registration', level=2)
    for run in reg_heading.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x44, 0x91)

    reg_text = doc.add_paragraph(
        'Early bird tickets available until May 15, 2025. Limited seats available!'
    )
    reg_text.paragraph_format.space_after = Pt(6)

    price_para = doc.add_paragraph()
    price_para.add_run('General Admission: ').bold = True
    price_para.add_run('$299')

    early_para = doc.add_paragraph()
    early_para.add_run('Early Bird (before May 15): ').bold = True
    early_para.add_run('$199')

    student_para = doc.add_paragraph()
    student_para.add_run('Student/Nonprofit: ').bold = True
    student_para.add_run('$99 (with valid ID)')

    doc.add_paragraph()

    # Contact info
    contact_para = doc.add_paragraph()
    contact_para.add_run('For information: ').bold = True
    contact_para.add_run('events@techsummit2025.com | +1 (415) 555-0192')
    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(DOC_OUTPUT)
    print(f'Initial document created: {DOC_OUTPUT}')


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)
    create_banner_image()
    create_initial_document()

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DOC_OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
