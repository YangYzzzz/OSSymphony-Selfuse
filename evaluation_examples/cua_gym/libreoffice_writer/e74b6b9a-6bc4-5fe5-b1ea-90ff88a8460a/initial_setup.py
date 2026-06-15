"""
Initial Setup: Weekly schedule document with five day sections
Task ID: writer_txtfmt_061
Domain: libreoffice_writer

Creates /home/user/Desktop/weekly_schedule.docx with:
- Five sections each starting with a day name (Monday-Friday)
- Followed by a realistic list of activities per day
- All text in 12pt Liberation Sans regular black, no background
- Day names are plain (NOT bold, NOT colored, NO background) — agent must apply formatting
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'weekly_schedule'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_run_font(run, font_name='Liberation Sans', size_pt=12):
    """Apply base font settings to a run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = False
    run.italic = False
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Remove default empty paragraph if present
    # (Document() adds one by default, we'll use it for first content)

    schedule = [
        {
            'day': 'Monday',
            'activities': [
                '9:00 AM - Team standup meeting in Conference Room B',
                '10:30 AM - Review Q1 budget projections with Finance',
                '12:00 PM - Lunch with client Sarah Mitchell',
                '2:00 PM - Code review for payment module refactor',
                '4:00 PM - Weekly one-on-one with manager',
                '5:30 PM - Gym session at LA Fitness',
            ],
        },
        {
            'day': 'Tuesday',
            'activities': [
                '8:30 AM - Prepare slides for board presentation',
                '10:00 AM - Product roadmap planning session',
                '11:30 AM - Customer support escalation review',
                '1:00 PM - Working lunch - UX research synthesis',
                '3:00 PM - Interview candidate for senior engineer role',
                '5:00 PM - Respond to pending client emails',
            ],
        },
        {
            'day': 'Wednesday',
            'activities': [
                '9:00 AM - All-hands company meeting (virtual)',
                '11:00 AM - Sprint retrospective with engineering team',
                '12:30 PM - Lunch break - visit farmers market',
                '2:00 PM - Architecture discussion for new microservices',
                '3:30 PM - Training session: AWS certification prep',
                '6:00 PM - Book club meeting at community center',
            ],
        },
        {
            'day': 'Thursday',
            'activities': [
                '9:30 AM - Marketing campaign strategy review',
                '10:30 AM - Legal review of vendor contracts',
                '12:00 PM - Department lunch celebration (David\'s promotion)',
                '2:00 PM - Deep work block: feature development',
                '4:30 PM - Bi-weekly metrics review with product manager',
                '7:00 PM - Yoga class at downtown studio',
            ],
        },
        {
            'day': 'Friday',
            'activities': [
                '9:00 AM - Weekly planning and priority setting',
                '10:00 AM - Cross-functional sync with design and DevOps',
                '11:30 AM - Demo new features to stakeholders',
                '1:00 PM - Team lunch at Rosario\'s Italian Kitchen',
                '3:00 PM - Code deployment and production monitoring',
                '4:30 PM - End-of-week review and next week prep',
            ],
        },
    ]

    for section in schedule:
        # Day name paragraph — plain, 12pt, no bold, no color, no background
        day_para = doc.add_paragraph()
        day_run = day_para.add_run(section['day'])
        set_run_font(day_run, size_pt=12)

        # Activities
        for activity in section['activities']:
            act_para = doc.add_paragraph()
            act_run = act_para.add_run(activity)
            set_run_font(act_run, size_pt=12)

        # Blank line between sections (except after last)
        if section['day'] != 'Friday':
            spacer = doc.add_paragraph()
            spacer.add_run('')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
