"""
Reward Script: Mail merge letter with embedded benefit elections table
Task ID: writer_mt_030
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Greeting contains EmployeeName merge field
  Component 2 (0.3): Document has a 4-row x 2-col table
  Component 3 (0.2): Table column 1 has correct benefit category labels
  Component 4 (0.2): Table column 2 has correct merge fields
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_030'


def persist_app_state(domain):
    """Save any unsaved changes in LibreOffice Writer."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_writer", "libreoffice_calc", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify mail merge letter with benefit elections table.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Greeting contains EmployeeName merge field (0.3 points)
    # Initial has "Dear Employee," — golden has "Dear <<EmployeeName>>,"
    try:
        greeting_has_merge_field = any(
            re.search(r'[«<{].*EmployeeName.*[»>}]', para.text, re.IGNORECASE)
            for para in doc.paragraphs
            if para.text.strip().lower().startswith('dear') and 'employeename' in para.text.lower()
        )
        if greeting_has_merge_field:
            print(f"PASS: Component 1 — Greeting contains EmployeeName merge field (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — No greeting with EmployeeName merge field found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Document contains a table with 4 rows and 2 columns (0.3 points)
    # Initial has 0 tables — golden has 1 table with 4 rows x 2 cols
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 4 and num_cols == 2:
                print(f"PASS: Component 2 — Table found with 4 rows x 2 cols (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Table has {num_rows} rows x {num_cols} cols, expected 4x2")
        else:
            print(f"FAIL: Component 2 — No tables found in document (expected 1)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Table column 1 has correct benefit category labels (0.2 points)
    # Expected labels: Health Plan, Dental Plan, Vision Plan, 401(k) Contribution
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            expected_labels = ['health plan', 'dental plan', 'vision plan', '401(k) contribution']
            actual_labels = []
            for ri in range(min(len(table.rows), 4)):
                cell_text = table.cell(ri, 0).text.strip().lower()
                actual_labels.append(cell_text)

            matches = 0
            for expected in expected_labels:
                if any(expected in actual for actual in actual_labels):
                    matches += 1

            if matches == 4:
                print(f"PASS: Component 3 — All 4 benefit labels found in column 1 (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — Only {matches}/4 labels matched. Actual: {actual_labels}")
        else:
            print(f"FAIL: Component 3 — No table to check labels")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Table column 2 has merge fields for benefit data (0.2 points)
    # Expected merge fields: HealthPlan, DentalPlan, VisionPlan, 401kContribution
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            expected_fields = ['healthplan', 'dentalplan', 'visionplan', '401kcontribution']
            found_fields = 0
            for ri in range(min(len(table.rows), 4)):
                cell_text = table.cell(ri, 1).text.strip().lower()
                # Check for merge field markers (guillemets, angle brackets, curly braces)
                for field in expected_fields:
                    if field in cell_text:
                        found_fields += 1
                        break

            if found_fields == 4:
                print(f"PASS: Component 4 — All 4 merge fields found in column 2 (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 4 — Only {found_fields}/4 merge fields matched")
        else:
            print(f"FAIL: Component 4 — No table to check merge fields")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
