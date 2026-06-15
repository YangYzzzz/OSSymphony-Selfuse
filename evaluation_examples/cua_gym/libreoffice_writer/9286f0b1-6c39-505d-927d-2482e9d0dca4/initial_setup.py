"""
Initial Setup: Generic benefits confirmation letter (no merge fields, no table)
Task ID: writer_mt_030
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_030'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Company letterhead ---
    header_para = doc.add_paragraph()
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header_para.add_run("Meridian Healthcare Solutions")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8E)

    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_sub = sub_header.add_run("Human Resources Department")
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    addr = doc.add_paragraph()
    addr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_addr = addr.add_run("4500 Innovation Drive, Suite 200  |  Austin, TX 78701  |  (512) 555-0147")
    run_addr.font.size = Pt(9)
    run_addr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Horizontal rule via bottom border on an empty paragraph
    doc.add_paragraph()  # spacer

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(12)
    run_date = date_para.add_run("March 15, 2026")
    run_date.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # --- Greeting (generic, no merge field) ---
    greeting = doc.add_paragraph()
    run_greet = greeting.add_run("Dear Employee,")
    run_greet.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # --- Body paragraphs ---
    body1 = doc.add_paragraph()
    run_b1 = body1.add_run(
        "Thank you for completing your annual benefits enrollment for the 2026 plan year. "
        "This letter serves as confirmation of the benefit elections you have made during "
        "the open enrollment period. Please review the information below carefully and "
        "contact Human Resources if you notice any discrepancies."
    )
    run_b1.font.size = Pt(11)

    body2 = doc.add_paragraph()
    run_b2 = body2.add_run(
        "Your selected benefits will take effect on April 1, 2026. You will receive "
        "your updated insurance cards and enrollment materials from each provider within "
        "two to three weeks of the effective date. If you do not receive your materials "
        "by April 21, 2026, please contact the Benefits Administration team at "
        "benefits@meridianhcs.com or extension 4200."
    )
    run_b2.font.size = Pt(11)

    body3 = doc.add_paragraph()
    run_b3 = body3.add_run(
        "Please note that changes to your elections can only be made during the annual "
        "open enrollment period or within 30 days of a qualifying life event, such as "
        "marriage, birth of a child, or loss of other coverage. Documentation of the "
        "qualifying event will be required."
    )
    run_b3.font.size = Pt(11)

    body4 = doc.add_paragraph()
    run_b4 = body4.add_run(
        "We value your contributions to Meridian Healthcare Solutions and are committed "
        "to providing a comprehensive benefits package that supports your health and "
        "financial well-being. Should you have any questions regarding your benefits, "
        "please do not hesitate to reach out to the HR Benefits team."
    )
    run_b4.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # --- Closing ---
    closing = doc.add_paragraph()
    run_cl = closing.add_run("Sincerely,")
    run_cl.font.size = Pt(11)

    doc.add_paragraph()  # spacer for signature

    sig_name = doc.add_paragraph()
    run_sig = sig_name.add_run("Patricia Alvarez")
    run_sig.bold = True
    run_sig.font.size = Pt(11)

    sig_title = doc.add_paragraph()
    run_title = sig_title.add_run("Director of Human Resources")
    run_title.font.size = Pt(11)

    sig_company = doc.add_paragraph()
    run_comp = sig_company.add_run("Meridian Healthcare Solutions")
    run_comp.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
