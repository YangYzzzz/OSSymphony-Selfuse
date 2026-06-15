"""
Initial Setup: Create a multi-page portrait Writer document with a wide financial
comparison table on page 4 that gets cut off.
Task ID: writer_biz_040
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_040'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_body_paragraph(doc, text, bold_prefix=None):
    para = doc.add_paragraph()
    if bold_prefix:
        run_b = para.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return para


def create_initial():
    doc = Document()

    # Set default page layout: portrait, letter size
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ===================== PAGE 1 =====================
    doc.add_heading('Meridian Technologies Inc.', level=0)
    doc.add_heading('Annual Financial Review 2025', level=1)

    add_body_paragraph(doc,
        'This report presents a comprehensive overview of Meridian Technologies\' '
        'financial performance across all operating divisions for the fiscal year ending '
        'December 31, 2025. The analysis covers revenue growth, operating margins, capital '
        'expenditure, and strategic investment returns.')

    doc.add_heading('Executive Summary', level=2)
    add_body_paragraph(doc,
        'Meridian Technologies achieved record revenue of $4.23 billion in FY2025, '
        'representing a 14.7% year-over-year increase. Operating margins expanded to '
        '22.3%, driven primarily by strong performance in the Cloud Infrastructure and '
        'Enterprise Security divisions.')

    add_body_paragraph(doc,
        'Key highlights include the successful launch of the Aurora platform, which '
        'generated $312 million in first-year revenue, and the acquisition of DataStream '
        'Analytics for $890 million, which has been fully integrated into our Data '
        'Intelligence division.')

    doc.add_heading('Market Context', level=2)
    add_body_paragraph(doc,
        'The global enterprise technology market grew by 8.2% in 2025, reaching $1.4 '
        'trillion. Within our addressable segments, cloud infrastructure spending increased '
        '19.1%, cybersecurity budgets rose 16.4%, and data analytics investments grew 21.7%. '
        'Meridian outperformed each of these benchmarks.')

    add_body_paragraph(doc,
        'Despite macroeconomic headwinds including elevated interest rates and supply chain '
        'disruptions in semiconductor procurement, our diversified portfolio and strategic '
        'long-term contracts provided resilience. Customer retention remained strong at '
        '94.6%, with net revenue retention reaching 118%.')

    # ===================== PAGE 2 =====================
    doc.add_page_break()
    doc.add_heading('Division Performance Overview', level=1)

    doc.add_heading('Cloud Infrastructure Division', level=2)
    add_body_paragraph(doc,
        'The Cloud Infrastructure Division generated $1.47 billion in revenue, a 21.3% '
        'increase from the prior year. This growth was fueled by enterprise migration '
        'projects, with 142 new enterprise contracts signed during the year. The Aurora '
        'platform\'s multi-cloud orchestration capabilities were a key differentiator.')

    add_body_paragraph(doc,
        'Operating margin for the division reached 26.8%, up from 23.1% in FY2024, '
        'reflecting economies of scale in data center operations and improved utilization '
        'rates across our 14 global facilities.')

    doc.add_heading('Enterprise Security Division', level=2)
    add_body_paragraph(doc,
        'Enterprise Security posted revenue of $987 million, growing 18.6% year-over-year. '
        'The division benefited from increased regulatory compliance requirements and '
        'rising cyber threat sophistication. Our Zero Trust architecture solution saw '
        'adoption by 67 Fortune 500 companies.')

    doc.add_heading('Data Intelligence Division', level=2)
    add_body_paragraph(doc,
        'Following the DataStream Analytics acquisition, the Data Intelligence Division '
        'reached $892 million in combined revenue. Integration synergies exceeded initial '
        'projections by 15%, with cross-selling driving $78 million in incremental revenue.')

    doc.add_heading('Professional Services Division', level=2)
    add_body_paragraph(doc,
        'Professional Services contributed $883 million, with consulting engagement revenue '
        'growing 11.2%. The division maintained a utilization rate of 82.4% and expanded '
        'its delivery team to 3,200 consultants across 28 countries.')

    # ===================== PAGE 3 =====================
    doc.add_page_break()
    doc.add_heading('Regional Performance', level=1)

    add_body_paragraph(doc,
        'Meridian Technologies operates across four primary geographic regions: North '
        'America, Europe Middle East and Africa (EMEA), Asia-Pacific (APAC), and Latin '
        'America. Each region demonstrated positive growth trajectories, though with '
        'varying dynamics.')

    doc.add_heading('North America', level=2)
    add_body_paragraph(doc,
        'North America remains our largest market at $2.28 billion (53.9% of total revenue), '
        'growing 12.4%. Federal government contracts contributed $340 million, with notable '
        'wins including the Department of Defense cybersecurity modernization program valued '
        'at $127 million over three years.')

    doc.add_heading('EMEA', level=2)
    add_body_paragraph(doc,
        'EMEA revenue reached $1.06 billion (25.1% of total), growing 16.8%. The UK and '
        'Germany were the strongest performers, with France and the Nordics showing '
        'accelerating adoption of our cloud platform.')

    doc.add_heading('Asia-Pacific', level=2)
    add_body_paragraph(doc,
        'APAC delivered $672 million (15.9% of total), growing 19.3%. Japan and Australia '
        'led growth, while our expanded presence in India and Southeast Asia began '
        'generating meaningful enterprise traction.')

    doc.add_heading('Latin America', level=2)
    add_body_paragraph(doc,
        'Latin America contributed $218 million (5.1% of total), growing 22.1%. Brazil '
        'and Mexico drove the majority of regional revenue, with cloud infrastructure '
        'demand accelerating as digital transformation initiatives gained momentum.')

    add_body_paragraph(doc,
        'Our global delivery model, combining regional sales teams with centralized product '
        'development, continues to provide competitive advantages in cost efficiency and '
        'speed to market across all regions.')

    # ===================== PAGE 4 — WIDE TABLE =====================
    doc.add_page_break()
    doc.add_heading('Financial Comparison Table — All Divisions by Quarter', level=1)

    add_body_paragraph(doc,
        'The following table provides a detailed quarterly breakdown of key financial '
        'metrics across all operating divisions. This comprehensive view enables '
        'quarter-over-quarter trend analysis and cross-divisional performance comparison.')

    # Wide table with many columns — designed to be cut off in portrait
    headers = [
        'Division', 'Q1 Revenue ($M)', 'Q1 COGS ($M)', 'Q1 Gross Margin',
        'Q2 Revenue ($M)', 'Q2 COGS ($M)', 'Q2 Gross Margin',
        'Q3 Revenue ($M)', 'Q3 COGS ($M)', 'Q3 Gross Margin',
        'Q4 Revenue ($M)', 'Q4 COGS ($M)', 'Q4 Gross Margin',
        'FY Total ($M)', 'YoY Growth'
    ]

    data_rows = [
        ['Cloud Infrastructure', '327.4', '239.6', '26.9%',
         '348.1', '254.8', '26.8%',
         '381.6', '278.2', '27.1%',
         '413.2', '302.9', '26.7%',
         '1,470.3', '21.3%'],
        ['Enterprise Security', '218.9', '162.0', '26.0%',
         '237.5', '174.1', '26.7%',
         '256.8', '186.4', '27.4%',
         '273.4', '196.8', '28.0%',
         '986.6', '18.6%'],
        ['Data Intelligence', '198.7', '153.6', '22.7%',
         '212.4', '162.5', '23.5%',
         '231.6', '175.4', '24.3%',
         '249.3', '186.2', '25.3%',
         '892.0', '34.2%'],
        ['Professional Services', '203.1', '158.4', '22.0%',
         '211.8', '163.1', '23.0%',
         '224.7', '171.6', '23.6%',
         '243.5', '184.9', '24.1%',
         '883.1', '11.2%'],
        ['Corporate & Other', '(12.4)', '(8.1)', 'N/A',
         '(11.8)', '(7.6)', 'N/A',
         '(13.2)', '(8.9)', 'N/A',
         '(14.1)', '(9.7)', 'N/A',
         '(51.5)', 'N/A'],
        ['CONSOLIDATED', '935.7', '705.5', '24.6%',
         '998.0', '746.9', '25.2%',
         '1,081.5', '802.7', '25.8%',
         '1,165.3', '861.1', '26.1%',
         '4,180.5', '14.7%'],
    ]

    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = 'Table Grid'

    # Headers
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

    # Data
    for i, row_data in enumerate(data_rows):
        for j, val in enumerate(row_data):
            cell = table.cell(i + 1, j)
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            if i == len(data_rows) - 1:  # CONSOLIDATED row
                run.bold = True

    # ===================== PAGE 5 =====================
    doc.add_page_break()
    doc.add_heading('Strategic Outlook and Recommendations', level=1)

    add_body_paragraph(doc,
        'Based on the financial performance reviewed above, management recommends the '
        'following strategic priorities for FY2026:')

    priorities = [
        'Accelerate Aurora platform adoption with targeted enterprise migration incentives '
        'and expanded partner ecosystem certification programs.',
        'Invest $450 million in next-generation AI-driven security analytics capabilities, '
        'building on the Zero Trust architecture momentum.',
        'Expand APAC operations with new data center facilities in Singapore and Mumbai, '
        'targeting $900 million in regional revenue by FY2027.',
        'Pursue strategic acquisitions in the edge computing space to complement existing '
        'cloud infrastructure capabilities.',
        'Increase R&D spending to 16% of revenue (from 14.2%) to maintain technology '
        'leadership and accelerate product innovation cycles.',
    ]
    for p in priorities:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('Capital Allocation', level=2)
    add_body_paragraph(doc,
        'The Board has approved a capital allocation framework for FY2026 totaling $2.1 '
        'billion, distributed across organic growth investments ($1.2B), strategic M&A '
        '($600M), shareholder returns via dividends and buybacks ($200M), and debt '
        'reduction ($100M). This framework balances growth ambitions with prudent '
        'financial management.')

    add_body_paragraph(doc,
        'The company maintains a strong balance sheet with $3.4 billion in cash and '
        'short-term investments, a debt-to-equity ratio of 0.42, and an investment-grade '
        'credit rating of A- from Standard and Poor\'s.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
