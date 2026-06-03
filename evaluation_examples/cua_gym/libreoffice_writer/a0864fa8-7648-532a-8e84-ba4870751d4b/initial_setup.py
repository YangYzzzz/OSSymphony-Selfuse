"""
Initial Setup: Short story collection document with 3 chapters flowing continuously (no page breaks)
Task ID: writer_page_039
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'short_story_collection'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup: A4 portrait, 2.54cm margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # --- Chapter 1: The Beginning ---
    h1 = doc.add_heading('Chapter 1: The Beginning', level=1)
    # NOTE: NO page_break_before on Chapter 1 (it's the first chapter)

    doc.add_paragraph(
        "The morning light crept through the curtains of the old farmhouse as Elena "
        "stirred from her restless sleep. Outside, the rolling hills of Millbrook Valley "
        "were draped in a thin veil of mist, the kind that only came in early September "
        "when summer reluctantly surrendered to autumn."
    )
    doc.add_paragraph(
        "She dressed quietly, careful not to wake her grandmother, and slipped out the "
        "back door into the garden. The roses her grandfather had planted decades ago "
        "still bloomed along the stone wall, their crimson petals holding beads of dew "
        "that sparkled in the growing daylight."
    )
    doc.add_paragraph(
        "Elena had returned to Millbrook for the first time in eleven years. The letter "
        "from her grandmother's solicitor had arrived on a Thursday, brief and formal: "
        "'Your presence is requested at Hawthorn Farm regarding matters of the estate.' "
        "She had not expected to feel anything. Yet standing in the familiar garden, "
        "grief settled over her like a second skin."
    )
    doc.add_paragraph(
        "The old oak tree at the far end of the garden still bore the rope of the swing "
        "her father had hung for her when she was six. It was frayed now, dangling from "
        "the branch like a forgotten promise. She walked toward it slowly, her footsteps "
        "soft on the wet grass, and placed one hand against the rough bark."
    )
    doc.add_paragraph(
        "'You came back,' said a voice behind her. Elena turned to find her grandmother "
        "standing at the garden gate, wrapped in a wool shawl despite the mildness of "
        "the morning. Margaret Hartley had always risen before dawn. Some things, it "
        "seemed, did not change."
    )

    # --- Chapter 2: The Journey ---
    # IMPORTANT: NO page break here in initial state (that's the task!)
    h2 = doc.add_heading('Chapter 2: The Journey', level=1)

    doc.add_paragraph(
        "Three days after her arrival, Elena found the old road atlas in the bureau "
        "drawer of her grandmother's study. It was a battered thing, its spine cracked "
        "and several pages dog-eared, covered in pencilled annotations in her grandfather's "
        "precise handwriting. Routes circled, small crosses marking towns of significance "
        "whose meaning she could not fathom."
    )
    doc.add_paragraph(
        "Beneath the atlas lay a packet of letters bound with garden twine. The envelopes "
        "were yellowed, each addressed in the same careful hand to 'W. Hartley, Poste "
        "Restante, Lisbonne.' Elena turned them over in her hands, reading the postmarks: "
        "1962, 1963, 1964. Her grandfather had never spoken of Portugal."
    )
    doc.add_paragraph(
        "She borrowed her grandmother's aging Volvo and set out along the valley road "
        "toward the village of Ashford, where the county archives were housed in a "
        "converted mill building beside the river. The archivist, a cheerful woman named "
        "Mrs. Devereux, was delighted to assist. 'The Hartley family?' she said. 'Oh yes, "
        "they have quite a history in this part of the world.'"
    )
    doc.add_paragraph(
        "The documents Mrs. Devereux produced were a revelation. William Hartley, Elena's "
        "grandfather, had served as a civil engineer on a dam project in the Alentejo "
        "region of Portugal in the early 1960s. He had worked there for three years before "
        "returning to England to marry Margaret Sayers in 1965. The letters, Elena now "
        "understood, were written during that time — but not by her grandfather."
    )
    doc.add_paragraph(
        "Elena drove back to Hawthorn Farm in the fading afternoon light, the packet of "
        "letters on the passenger seat beside her. Her grandmother was sitting in the "
        "kitchen when she arrived, a pot of tea already brewing, as though she had known "
        "exactly when her granddaughter would return."
    )

    # --- Chapter 3: The Return ---
    # IMPORTANT: NO page break here in initial state (that's the task!)
    h3 = doc.add_heading('Chapter 3: The Return', level=1)

    doc.add_paragraph(
        "It took Margaret Hartley two days before she was ready to speak. She sat with "
        "Elena at the kitchen table on a Sunday morning, the letters spread between them "
        "like a map of a country neither of them had visited. Outside, rain tapped steadily "
        "against the windowpanes."
    )
    doc.add_paragraph(
        "'Her name was Inês,' Margaret said at last, her voice steady, matter-of-fact in "
        "the way of a woman who has long since made her peace with something. 'She was a "
        "schoolteacher in the town near the dam site. They met at a village festival in "
        "the autumn of 1962. She wrote to him for three years after he came home.'"
    )
    doc.add_paragraph(
        "Elena looked at her grandmother carefully. 'Did you know, when you married him?' "
        "Margaret was quiet for a moment. 'I knew there had been someone,' she said. "
        "'He told me that much. He never told me her name. I suppose I never truly wanted "
        "to know it.' She smoothed one of the envelopes with her fingertips. 'He kept "
        "the letters all this time. That tells its own story.'"
    )
    doc.add_paragraph(
        "Elena spent the rest of her visit helping her grandmother sort through the "
        "accumulated belongings of sixty years of marriage. They worked methodically, "
        "room by room, deciding what to keep, what to give away, what to let go. It was "
        "quiet work, companionable in its silence. On the last morning, Elena packed her "
        "bag and stood in the hallway of Hawthorn Farm for what she suspected might be "
        "the final time."
    )
    doc.add_paragraph(
        "Margaret walked her out to the car. The mist had returned to the valley, softer "
        "now than it had been on the morning Elena arrived. They embraced at the gate, "
        "and Elena held on a little longer than she meant to. As she drove away down the "
        "lane, she watched in the rearview mirror until the old farmhouse disappeared "
        "behind the hedgerows. She did not look away until it was entirely gone."
    )

    # Ensure Desktop directory exists (create it if needed)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the initial artifact in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
