"""
Initial Setup: Create a Writer document with default styles and technical documentation content.
Task ID: writer_tech_084
Domain: libreoffice_writer

The document has realistic technical documentation content using only default styles.
No custom styles exist yet - the agent must create the style hierarchy.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_084'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    doc.add_heading("CloudSync API Integration Guide", level=0)

    # Introduction
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "CloudSync is a distributed file synchronization platform designed for "
        "enterprise teams managing large-scale data pipelines. This guide covers "
        "the RESTful API endpoints, authentication flows, and webhook configuration "
        "required to integrate CloudSync into your existing infrastructure."
    )
    doc.add_paragraph(
        "The API supports OAuth 2.0 bearer tokens with configurable scopes. All "
        "requests must include a valid Authorization header. Rate limits are enforced "
        "at 500 requests per minute per API key."
    )

    # Authentication section
    doc.add_heading("2. Authentication", level=1)
    doc.add_paragraph(
        "Before making API calls, you must register your application in the "
        "CloudSync Developer Portal at https://developers.cloudsync.io. Upon "
        "registration, you will receive a client_id and client_secret pair."
    )
    doc.add_paragraph(
        "Token exchange is performed via a POST request to /oauth/token with "
        "grant_type=client_credentials. Tokens expire after 3600 seconds by default, "
        "but this can be configured in your application settings."
    )
    doc.add_paragraph(
        "For server-to-server integrations, we recommend storing credentials in "
        "environment variables rather than configuration files. Never commit API "
        "keys to version control systems."
    )

    # Endpoints section
    doc.add_heading("3. Core Endpoints", level=1)

    doc.add_heading("3.1 File Operations", level=2)
    doc.add_paragraph(
        "The /v2/files endpoint supports CRUD operations on synchronized files. "
        "Upload files using multipart/form-data encoding with a maximum payload "
        "size of 500 MB per request. Chunked uploads are available for files "
        "exceeding this limit via the /v2/files/chunked endpoint."
    )

    # Add a table of endpoints
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ["Method", "Endpoint", "Description"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    endpoints = [
        ["GET", "/v2/files", "List all synced files with pagination"],
        ["POST", "/v2/files", "Upload a new file to the sync pool"],
        ["GET", "/v2/files/{id}", "Retrieve metadata for a specific file"],
        ["PUT", "/v2/files/{id}", "Update file content or metadata"],
        ["DELETE", "/v2/files/{id}", "Remove a file from synchronization"],
    ]
    for r, row_data in enumerate(endpoints, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_heading("3.2 Webhook Configuration", level=2)
    doc.add_paragraph(
        "Webhooks allow your application to receive real-time notifications when "
        "files are created, modified, or deleted. Register webhook URLs via the "
        "/v2/webhooks endpoint. Each webhook must specify an event type and a "
        "target URL that accepts POST requests."
    )
    doc.add_paragraph(
        "Webhook payloads are signed using HMAC-SHA256 with your webhook secret. "
        "Always verify the X-CloudSync-Signature header before processing events "
        "to prevent spoofing attacks."
    )

    # Error handling section
    doc.add_heading("4. Error Handling", level=1)
    doc.add_paragraph(
        "All API responses follow standard HTTP status codes. Successful operations "
        "return 2xx codes, client errors return 4xx, and server errors return 5xx. "
        "Error responses include a JSON body with 'error_code' and 'message' fields."
    )
    doc.add_paragraph(
        "When receiving a 429 Too Many Requests response, implement exponential "
        "backoff with a base delay of 1 second. The Retry-After header indicates "
        "the minimum wait time before retrying."
    )

    # Best practices
    doc.add_heading("5. Best Practices", level=1)
    doc.add_paragraph(
        "Use connection pooling for high-throughput integrations. The recommended "
        "pool size is 10 connections per worker thread. Enable gzip compression "
        "by setting Accept-Encoding: gzip in request headers to reduce bandwidth "
        "consumption by approximately 60-70%."
    )
    doc.add_paragraph(
        "Implement idempotency keys for all write operations. Pass a unique "
        "X-Idempotency-Key header with each POST or PUT request to ensure safe "
        "retries without duplicate side effects."
    )
    doc.add_paragraph(
        "Monitor your integration health using the /v2/status endpoint, which "
        "returns current rate limit usage, webhook delivery statistics, and "
        "synchronization lag metrics."
    )

    # Appendix
    doc.add_heading("Appendix A: Status Codes", level=1)
    status_table = doc.add_table(rows=7, cols=2)
    status_table.style = "Table Grid"
    status_table.cell(0, 0).text = "Code"
    status_table.cell(0, 1).text = "Meaning"
    codes = [
        ["200", "Success"],
        ["201", "Resource created"],
        ["400", "Bad request - check parameters"],
        ["401", "Authentication required"],
        ["429", "Rate limit exceeded"],
        ["500", "Internal server error"],
    ]
    for r, row_data in enumerate(codes, 1):
        for c, val in enumerate(row_data):
            status_table.cell(r, c).text = val

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
