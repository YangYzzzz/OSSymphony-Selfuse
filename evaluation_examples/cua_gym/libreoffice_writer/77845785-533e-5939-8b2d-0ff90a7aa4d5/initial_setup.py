"""
Initial Setup: Create a document with two code examples in Default Paragraph Style
Task ID: writer_tech_026
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_026'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    heading = doc.add_heading("Python Data Processing Guide", level=1)

    # --- Intro paragraph ---
    intro = doc.add_paragraph(
        "This guide demonstrates common data processing patterns in Python. "
        "The following examples show how to read CSV files and perform basic "
        "statistical analysis on datasets."
    )

    # --- Section 1 ---
    doc.add_heading("Reading and Filtering Data", level=2)

    p1 = doc.add_paragraph(
        "The first step in any data pipeline is loading the raw data from disk. "
        "The example below reads a CSV file containing quarterly sales records "
        "and filters out entries below a revenue threshold."
    )

    # --- Code Block 1 (Default Paragraph Style, Liberation Mono font, no border/bg) ---
    code1_text = (
        'import csv\n'
        '\n'
        'def load_sales_data(filepath, min_revenue=5000):\n'
        '    """Load sales CSV and filter by minimum revenue."""\n'
        '    results = []\n'
        '    with open(filepath, "r") as f:\n'
        '        reader = csv.DictReader(f)\n'
        '        for row in reader:\n'
        '            if float(row["revenue"]) >= min_revenue:\n'
        '                results.append(row)\n'
        '    return results\n'
        '\n'
        'sales = load_sales_data("quarterly_sales_2025.csv")\n'
        'print(f"Filtered records: {len(sales)}")'
    )

    code1_para = doc.add_paragraph()
    code1_run = code1_para.add_run(code1_text)
    code1_run.font.name = "Liberation Mono"
    code1_run.font.size = Pt(9)

    # --- Explanation paragraph ---
    p2 = doc.add_paragraph(
        "After loading and filtering, we typically need to compute summary statistics. "
        "The next example calculates the mean, median, and standard deviation of "
        "the revenue column from the filtered dataset."
    )

    # --- Section 2 ---
    doc.add_heading("Computing Summary Statistics", level=2)

    p3 = doc.add_paragraph(
        "Statistical summaries help identify trends and outliers in the data. "
        "Below is a utility function that computes key metrics without relying "
        "on external libraries like NumPy or pandas."
    )

    # --- Code Block 2 (Default Paragraph Style, Liberation Mono font, no border/bg) ---
    code2_text = (
        'def compute_stats(records, field="revenue"):\n'
        '    """Compute mean, median, and std dev for a numeric field."""\n'
        '    values = sorted(float(r[field]) for r in records)\n'
        '    n = len(values)\n'
        '    mean = sum(values) / n\n'
        '    median = values[n // 2] if n % 2 else (\n'
        '        values[n // 2 - 1] + values[n // 2]) / 2\n'
        '    variance = sum((v - mean) ** 2 for v in values) / n\n'
        '    std_dev = variance ** 0.5\n'
        '    return {"mean": mean, "median": median, "std_dev": std_dev}\n'
        '\n'
        'stats = compute_stats(sales)\n'
        'for key, val in stats.items():\n'
        '    print(f"{key}: {val:.2f}")'
    )

    code2_para = doc.add_paragraph()
    code2_run = code2_para.add_run(code2_text)
    code2_run.font.name = "Liberation Mono"
    code2_run.font.size = Pt(9)

    # --- Closing paragraph ---
    closing = doc.add_paragraph(
        "These patterns form the foundation of most data processing workflows. "
        "For larger datasets, consider using pandas DataFrames or database "
        "connections for improved performance and memory efficiency."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
