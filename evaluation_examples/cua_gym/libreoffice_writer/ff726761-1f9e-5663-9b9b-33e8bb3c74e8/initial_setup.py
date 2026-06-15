"""
Initial Setup: Corporate report with 4 Heading 1 paragraphs in default black color.
Task ID: osworld_writer_heading_styles_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_heading_styles_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Ensure Heading 1 style has NO custom color (uses default black/auto)
    # We will explicitly clear any color on Heading 1 style to ensure it's black
    heading1_style = doc.styles['Heading 1']
    # Remove any existing color from the style definition
    rPr = heading1_style.element.find(qn('w:rPr'))
    if rPr is not None:
        color_el = rPr.find(qn('w:color'))
        if color_el is not None:
            rPr.remove(color_el)

    # Section 1: Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'Meridian Solutions Group delivered strong performance in Q4 2024, '
        'achieving revenue of $4.82 billion — a 12.3% increase year-over-year. '
        'Operating margins improved to 18.7%, driven by efficiency gains in our '
        'North American and Asia-Pacific divisions. This report outlines the '
        'strategic highlights and forward-looking initiatives for fiscal year 2025.'
    )
    doc.add_paragraph(
        'Key performance indicators exceeded board targets across all three business '
        'units: Enterprise Software, Cloud Services, and Professional Consulting. '
        'Customer satisfaction scores reached an all-time high of 4.6 out of 5.0, '
        'reflecting the positive impact of our service modernization program.'
    )

    # Section 2: Financial Performance
    doc.add_heading('Financial Performance', level=1)
    doc.add_paragraph(
        'Total revenue for fiscal year 2024 was $18.4 billion, representing growth '
        'of 9.8% compared to $16.76 billion in 2023. Gross profit margin held steady '
        'at 41.2%, and EBITDA reached $3.44 billion. Capital expenditures of $820 '
        'million were allocated primarily to data center expansion and R&D facilities.'
    )
    doc.add_paragraph(
        'The Enterprise Software division generated $7.9 billion in revenue, up 14.1% '
        'from the prior year. Cloud Services contributed $6.2 billion (+8.6%), while '
        'Professional Consulting posted $4.3 billion (+4.3%). Free cash flow for the '
        'year was $2.61 billion, supporting continued investment in growth initiatives.'
    )

    # Add a table with financial data
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['Division', 'Revenue 2024 ($B)', 'Revenue 2023 ($B)', 'Growth (%)']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    rows_data = [
        ('Enterprise Software', '7.9', '6.92', '14.1%'),
        ('Cloud Services', '6.2', '5.71', '8.6%'),
        ('Professional Consulting', '4.3', '4.12', '4.3%'),
        ('Total', '18.4', '16.76', '9.8%'),
    ]
    for i, row_data in enumerate(rows_data, 1):
        for j, val in enumerate(row_data):
            table.cell(i, j).text = val

    doc.add_paragraph('')

    # Section 3: Strategic Initiatives
    doc.add_heading('Strategic Initiatives', level=1)
    doc.add_paragraph(
        'In 2025, Meridian Solutions Group will focus on four strategic pillars: '
        'AI-powered automation, geographic expansion into Southeast Asia, portfolio '
        'optimization through strategic divestitures, and talent development programs '
        'targeting 2,400 new hires across engineering and data science roles.'
    )
    doc.add_paragraph(
        'The AI Center of Excellence, established in Singapore in September 2024, '
        'has already deployed 17 automation workflows that collectively reduced '
        'operational costs by $145 million annually. Phase 2 of the program targets '
        '$380 million in savings by end of 2025 through intelligent process automation '
        'across our supply chain, customer service, and finance functions.'
    )
    doc.add_paragraph(
        'Geographic expansion into Vietnam, Indonesia, and the Philippines is on '
        'track, with three new regional offices scheduled to open by Q2 2025. '
        'Market analysis projects these markets will contribute $620 million in '
        'combined revenue by the end of fiscal year 2026.'
    )

    # Section 4: Outlook and Guidance
    doc.add_heading('Outlook and Guidance', level=1)
    doc.add_paragraph(
        'For fiscal year 2025, the board of directors has approved guidance of '
        '$20.1 – $20.6 billion in total revenue, reflecting projected growth of '
        '9.2% – 11.9%. Operating margins are expected to remain between 18.5% and '
        '19.2%, and earnings per share (diluted) guidance is set at $8.45 – $8.70.'
    )
    doc.add_paragraph(
        'Planned investments include $950 million in capital expenditures, with '
        '$400 million dedicated to cloud infrastructure upgrades and $310 million '
        'for research and development. The company reaffirms its commitment to '
        'returning value to shareholders through a quarterly dividend of $0.85 per '
        'share and a $1.5 billion share buyback program authorized through December 2025.'
    )
    doc.add_paragraph(
        'Management remains confident in the company\'s long-term trajectory and '
        'the resilience of our diversified business model. We will continue to '
        'monitor macroeconomic conditions, currency fluctuations, and regulatory '
        'developments across our operating regions, adjusting our strategy as needed '
        'to deliver sustainable value for all stakeholders.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
