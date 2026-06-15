"""
Initial Setup: Apply Heading 1 style to 'Executive Summary' title
Task ID: writer_biz_009
Domain: libreoffice_writer

Creates a business proposal document where 'Executive Summary' is the first
paragraph in Default Paragraph Style (Normal), bold, 14pt. The rest of the
document contains realistic business proposal content.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title: Executive Summary (Normal style, bold, 14pt) ---
    title_para = doc.add_paragraph()
    title_para.style = doc.styles['Normal']
    run = title_para.add_run('Executive Summary')
    run.bold = True
    run.font.size = Pt(14)

    # --- Executive Summary content ---
    doc.add_paragraph(
        'Meridian Consulting Group is pleased to present this comprehensive '
        'business proposal for the strategic partnership with Oakridge '
        'Technologies. This document outlines our approach to delivering '
        'an integrated digital transformation solution that will modernize '
        'Oakridge\'s customer-facing platforms and internal operations over '
        'the next 18 months.'
    )

    doc.add_paragraph(
        'Our proposed engagement encompasses three core workstreams: '
        'cloud infrastructure migration, customer experience redesign, '
        'and data analytics implementation. The total investment for this '
        'initiative is estimated at $2.4 million, with projected ROI of '
        '340% within the first three years of deployment.'
    )

    # --- Company Overview section ---
    overview_para = doc.add_paragraph()
    run = overview_para.add_run('Company Overview')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph(
        'Founded in 2012, Meridian Consulting Group has grown from a '
        'boutique advisory firm into a full-service technology consultancy '
        'with over 280 employees across offices in Chicago, Austin, and '
        'Seattle. We specialize in digital transformation for mid-market '
        'enterprises in the manufacturing, healthcare, and financial '
        'services sectors.'
    )

    doc.add_paragraph(
        'In the past fiscal year, Meridian completed 47 enterprise-level '
        'engagements with a 96% client satisfaction rating. Notable clients '
        'include Harmon Medical Systems, Lakeview Financial Corp, and '
        'Pacific Coast Manufacturing. Our team includes 35 certified cloud '
        'architects and 22 data science professionals.'
    )

    # --- Market Analysis section ---
    market_para = doc.add_paragraph()
    run = market_para.add_run('Market Analysis')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph(
        'The digital transformation market is projected to reach $3.4 '
        'trillion by 2026, with mid-market enterprises representing the '
        'fastest-growing segment at 28% CAGR. Oakridge Technologies '
        'operates in a competitive landscape where 73% of peer companies '
        'have already initiated some form of digital modernization.'
    )

    doc.add_paragraph(
        'Key industry drivers include rising customer expectations for '
        'seamless digital experiences, increasing regulatory requirements '
        'for data governance, and the growing need for real-time analytics '
        'capabilities. Organizations that delay transformation risk losing '
        '15-20% market share to digitally native competitors within '
        'the next five years.'
    )

    # --- Proposed Solution section ---
    solution_para = doc.add_paragraph()
    run = solution_para.add_run('Proposed Solution')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph(
        'Our three-phase approach begins with a comprehensive assessment '
        'of Oakridge\'s current technology stack and business processes. '
        'Phase 1 (months 1-4) focuses on cloud infrastructure migration '
        'using AWS and Azure hybrid architecture. Phase 2 (months 5-10) '
        'delivers the redesigned customer portal and mobile application. '
        'Phase 3 (months 11-18) implements the advanced analytics platform '
        'with machine learning capabilities for predictive insights.'
    )

    # --- Budget Summary section ---
    budget_para = doc.add_paragraph()
    run = budget_para.add_run('Budget Summary')
    run.bold = True
    run.font.size = Pt(13)

    # Budget table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Phase', 'Description', 'Investment']
    for col, h in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    data = [
        ['Phase 1', 'Cloud Infrastructure Migration', '$680,000'],
        ['Phase 2', 'Customer Experience Redesign', '$920,000'],
        ['Phase 3', 'Data Analytics Platform', '$800,000'],
        ['Total', '', '$2,400,000'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')  # spacing

    # --- Timeline section ---
    timeline_para = doc.add_paragraph()
    run = timeline_para.add_run('Timeline and Milestones')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph(
        'The project will commence on July 1, 2026, with final delivery '
        'targeted for December 31, 2027. Key milestones include the cloud '
        'migration completion by October 2026, customer portal beta launch '
        'in March 2027, and full analytics platform deployment by '
        'November 2027. Monthly progress reviews will be conducted with '
        'the Oakridge executive steering committee.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
