"""
Reward Script: Apply Heading 1 style to 'Executive Summary' title
Task ID: writer_biz_009
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): First paragraph has 'Heading 1' style applied
  Component 2 (0.4): First paragraph has 'Heading 1' style AND text is 'Executive Summary'
                      AND document structure is intact (paragraph count preserved)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_009'


def persist_app_state():
    """Save any unsaved LibreOffice Writer edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that the 'Executive Summary' paragraph has Heading 1 style applied.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 1 paragraph
    if len(doc.paragraphs) == 0:
        print("FAIL: Document has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    first_para = doc.paragraphs[0]
    first_style = first_para.style.name if first_para.style else None

    # Component 1: First paragraph has 'Heading 1' style (0.6 points)
    # This is the core task requirement. In initial_env the style is 'Normal'.
    try:
        if first_style == 'Heading 1':
            print(f"PASS: Component 1 - First paragraph style is 'Heading 1' (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 1 - Expected style 'Heading 1', found '{first_style}'")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: First paragraph has 'Heading 1' AND text is 'Executive Summary'
    #              AND document structure is intact (0.4 points)
    # Both sub-conditions are anchored to the style change: we only award points
    # if the style is Heading 1 (which fails on initial_env).
    try:
        if first_style == 'Heading 1':
            text_match = first_para.text.strip() == 'Executive Summary'
            para_count = len(doc.paragraphs)
            # Initial doc has 15 paragraphs; allow small tolerance (14-16)
            structure_ok = 14 <= para_count <= 16
            if text_match and structure_ok:
                print(f"PASS: Component 2 - Text is 'Executive Summary' and doc has {para_count} paragraphs (0.4 pts)")
                total_score += 0.4
            elif text_match:
                print(f"PARTIAL: Component 2 - Text matches but paragraph count is {para_count} (expected ~15) (0.2 pts)")
                total_score += 0.2
            elif structure_ok:
                print(f"PARTIAL: Component 2 - Structure intact ({para_count} paras) but text is '{first_para.text.strip()[:50]}' (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 2 - Text='{first_para.text.strip()[:50]}', paras={para_count}")
        else:
            print(f"FAIL: Component 2 - Skipped (style is not Heading 1)")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
