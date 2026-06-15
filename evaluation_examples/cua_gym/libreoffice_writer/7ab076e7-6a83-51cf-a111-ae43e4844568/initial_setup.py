"""
Initial Setup: Corporate Governance Report document
Task ID: writer_txtfmt_056
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'governance_report'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title paragraph ---
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('CORPORATE GOVERNANCE REPORT')
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    title_run.bold = False
    title_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # black

    # Blank line after title
    doc.add_paragraph()

    # --- Section 1: Board of Directors ---
    s1_head = doc.add_paragraph()
    s1_run = s1_head.add_run('1. Board of Directors')
    s1_run.font.name = 'Times New Roman'
    s1_run.font.size = Pt(13)
    s1_run.bold = True

    s1_body = doc.add_paragraph(
        'The Board of Directors consists of twelve members, seven of whom are independent '
        'non-executive directors. The Board held nine meetings during the fiscal year ending '
        'December 31, 2024. Key responsibilities include strategic oversight, risk management '
        'framework approval, and executive compensation review. The Chairman, Ms. Eleanor '
        'Whitfield, has served in this capacity since 2019 and brings extensive experience '
        'in financial services and regulatory compliance.'
    )
    s1_body.paragraph_format.space_after = Pt(8)

    # --- Section 2: Executive Compensation ---
    s2_head = doc.add_paragraph()
    s2_run = s2_head.add_run('2. Executive Compensation')
    s2_run.font.name = 'Times New Roman'
    s2_run.font.size = Pt(13)
    s2_run.bold = True

    s2_body = doc.add_paragraph(
        'Total compensation for the Chief Executive Officer, Mr. Jonathan Hargreaves, '
        'amounted to $4.2 million for fiscal year 2024, comprising a base salary of '
        '$1.1 million, annual performance bonus of $850,000, long-term incentive awards '
        'valued at $1.9 million, and benefits totaling $350,000. The Compensation Committee '
        'benchmarks executive pay against a peer group of 25 comparable companies in the '
        'sector. Variable pay constitutes 74% of total direct compensation.'
    )
    s2_body.paragraph_format.space_after = Pt(8)

    # --- Section 3: Audit Committee ---
    s3_head = doc.add_paragraph()
    s3_run = s3_head.add_run('3. Audit Committee')
    s3_run.font.name = 'Times New Roman'
    s3_run.font.size = Pt(13)
    s3_run.bold = True

    s3_body = doc.add_paragraph(
        'The Audit Committee comprises four independent directors, chaired by Mr. David '
        'Thornton, CPA. The Committee met eleven times during 2024 and oversees financial '
        'reporting integrity, internal controls, and the relationship with external auditors '
        'Pricewaterhouse & Partners LLP. The external audit fee for 2024 was $2.8 million. '
        'The Committee reviewed and approved the internal audit plan, monitored cybersecurity '
        'risks, and assessed the effectiveness of the internal control framework.'
    )
    s3_body.paragraph_format.space_after = Pt(8)

    # --- Section 4: Risk Management ---
    s4_head = doc.add_paragraph()
    s4_run = s4_head.add_run('4. Risk Management')
    s4_run.font.name = 'Times New Roman'
    s4_run.font.size = Pt(13)
    s4_run.bold = True

    s4_body = doc.add_paragraph(
        'Enterprise risk management is governed by the Risk Committee, which identifies, '
        'assesses, and monitors principal risks including market volatility, credit exposure, '
        'operational disruptions, and regulatory changes. The Company maintains a Risk '
        'Appetite Statement reviewed annually by the Board. Key risk indicators are reported '
        'monthly to the Risk Committee. In 2024, the Company invested $12.4 million in '
        'enhancing its cybersecurity infrastructure following a comprehensive third-party '
        'assessment conducted by TechSecure Advisory Group.'
    )
    s4_body.paragraph_format.space_after = Pt(8)

    # --- Section 5: Shareholder Rights ---
    s5_head = doc.add_paragraph()
    s5_run = s5_head.add_run('5. Shareholder Rights')
    s5_run.font.name = 'Times New Roman'
    s5_run.font.size = Pt(13)
    s5_run.bold = True

    s5_body = doc.add_paragraph(
        'The Company is committed to transparent and equitable treatment of all shareholders. '
        'Ordinary shareholders are entitled to one vote per share at the Annual General '
        'Meeting. The 2024 AGM was attended by shareholders representing 68.3% of issued '
        'share capital. Minority shareholder protections are embedded in the Articles of '
        'Association, including pre-emption rights on new share issuances. The Company '
        'provides quarterly earnings updates, investor roadshows, and maintains a dedicated '
        'investor relations portal at www.companyinvestors.com.'
    )
    s5_body.paragraph_format.space_after = Pt(8)

    # --- Section 6: Sustainability and ESG ---
    s6_head = doc.add_paragraph()
    s6_run = s6_head.add_run('6. Sustainability and ESG')
    s6_run.font.name = 'Times New Roman'
    s6_run.font.size = Pt(13)
    s6_run.bold = True

    s6_body = doc.add_paragraph(
        'Environmental, Social, and Governance (ESG) considerations are integrated into '
        'the Company\'s strategic planning and operational decisions. In 2024, the Company '
        'reduced its carbon emissions by 18% compared to the 2020 baseline and sourced '
        '42% of its energy from renewable sources. The Diversity and Inclusion Committee '
        'reports that women now represent 38% of senior leadership positions, up from 31% '
        'in 2021. The Company contributed $3.7 million to community development programs '
        'across fifteen countries in which it operates.'
    )
    s6_body.paragraph_format.space_after = Pt(8)

    # Ensure Desktop directory exists on VM
    os.makedirs(WORKDIR, exist_ok=True) if WORKDIR != '/home/user/Desktop' else None

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer with DISPLAY=:0
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
