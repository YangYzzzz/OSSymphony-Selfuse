"""
Initial Setup: Envelope document with delivery address in Beverly Hills, CA 90210
Task ID: writer_lec_047
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_047'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Configure envelope page (#10 envelope: 9.5 x 4.125 inches) ---
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(9.5)
    section.page_height = Inches(4.125)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # --- Return Address (upper left) ---
    return_para = doc.add_paragraph()
    return_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return_para.paragraph_format.space_after = Pt(0)
    run = return_para.add_run("Westfield Publishing Group")
    run.font.name = "Arial"
    run.font.size = Pt(10)

    return_line2 = doc.add_paragraph()
    return_line2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return_line2.paragraph_format.space_after = Pt(0)
    run = return_line2.add_run("4200 Wilshire Blvd, Suite 310")
    run.font.name = "Arial"
    run.font.size = Pt(10)

    return_line3 = doc.add_paragraph()
    return_line3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return_line3.paragraph_format.space_after = Pt(0)
    run = return_line3.add_run("Los Angeles, CA 90010")
    run.font.name = "Arial"
    run.font.size = Pt(10)

    # --- Spacer ---
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(36)
    spacer.paragraph_format.space_after = Pt(0)

    # --- Delivery Address (center area, slightly right) ---
    del_para1 = doc.add_paragraph()
    del_para1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    del_para1.paragraph_format.space_after = Pt(0)
    del_para1.paragraph_format.left_indent = Inches(2.0)
    run = del_para1.add_run("Dr. Helena Vasquez")
    run.font.name = "Arial"
    run.font.size = Pt(12)

    del_para2 = doc.add_paragraph()
    del_para2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    del_para2.paragraph_format.space_after = Pt(0)
    del_para2.paragraph_format.left_indent = Inches(2.0)
    run = del_para2.add_run("724 North Roxbury Drive")
    run.font.name = "Arial"
    run.font.size = Pt(12)

    del_para3 = doc.add_paragraph()
    del_para3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    del_para3.paragraph_format.space_after = Pt(0)
    del_para3.paragraph_format.left_indent = Inches(2.0)
    run = del_para3.add_run("Beverly Hills, CA 90210")
    run.font.name = "Arial"
    run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
