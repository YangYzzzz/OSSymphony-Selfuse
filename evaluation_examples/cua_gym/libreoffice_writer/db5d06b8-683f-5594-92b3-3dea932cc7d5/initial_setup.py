"""
Initial Setup: Apply 0.5cm first-line indent and justified alignment to body paragraphs in journal article
Task ID: writer_para_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_para_043'
OUTPUT = f'{WORKDIR}/Desktop/journal_article.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Heading 1 - title
    heading1 = doc.add_heading('Climate Change Adaptation Strategies in Coastal Cities', level=1)
    # (alignment left by default for headings - no changes)

    # Paragraph 2: Heading 2 - Abstract
    doc.add_heading('Abstract', level=2)

    # Paragraph 3: Body paragraph (NO first-line indent, NO justified alignment)
    para3 = doc.add_paragraph(
        'This paper examines the adaptation strategies employed by five major coastal cities '
        'in response to rising sea levels and increased storm intensity over the past decade.'
    )
    # Explicitly set left alignment (no first-line indent)
    para3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para3.paragraph_format.first_line_indent = None

    # Paragraph 4: Heading 2 - 1. Introduction
    doc.add_heading('1. Introduction', level=2)

    # Paragraph 5: Body paragraph (NO first-line indent, NO justified alignment)
    para5 = doc.add_paragraph(
        'Coastal cities are among the most vulnerable to the effects of climate change. '
        'Rising sea levels, combined with more frequent and intense storm events, pose '
        'significant threats to infrastructure and human safety.'
    )
    para5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para5.paragraph_format.first_line_indent = None

    # Paragraph 6: Body paragraph (NO first-line indent, NO justified alignment)
    para6 = doc.add_paragraph(
        'The economic impact is staggering. A recent World Bank report estimates that '
        'without intervention, annual flood damages to coastal cities could exceed '
        '$1 trillion by 2050.'
    )
    para6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para6.paragraph_format.first_line_indent = None

    # Paragraph 7: Heading 2 - 2. Study Areas
    doc.add_heading('2. Study Areas', level=2)

    # Paragraph 8: Body paragraph (NO first-line indent, NO justified alignment)
    para8 = doc.add_paragraph(
        'We selected five cities representing diverse geographic and economic contexts: '
        'Miami, Mumbai, Jakarta, Rotterdam, and Lagos. Each city faces unique challenges '
        'but shares the common threat of sea level rise.'
    )
    para8.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para8.paragraph_format.first_line_indent = None

    # Ensure the Desktop directory exists
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer with DISPLAY=:0
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
