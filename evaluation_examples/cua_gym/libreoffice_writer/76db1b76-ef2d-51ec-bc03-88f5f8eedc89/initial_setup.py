"""
Initial Setup: Technical manual with plain bold section/subsection titles (no Heading styles applied)
Task ID: osworld_writer_heading_styles_002
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_heading_styles_002'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_bold_heading(doc, text):
    """Add a paragraph with bold text styled as plain Normal (NOT Heading style)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return para


def create_initial():
    doc = Document()

    # Document title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('Network Infrastructure Management Guide')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # blank line

    intro = doc.add_paragraph(
        'This guide provides comprehensive instructions for managing and maintaining network '
        'infrastructure within enterprise environments. Follow the procedures outlined in each '
        'section to ensure optimal system performance and security compliance.'
    )

    doc.add_paragraph()  # blank line

    # ==========================================================================
    # SECTION 1 — plain bold, NOT Heading 1
    # ==========================================================================
    add_bold_heading(doc, 'Section 1: Network Architecture Overview')

    doc.add_paragraph(
        'The network architecture is designed to support high availability and fault tolerance '
        'across all critical business systems. The core infrastructure comprises three redundant '
        'data centers connected via dedicated fiber links operating at 10 Gbps.'
    )

    # Subsection 1.1 — plain bold, NOT Heading 2
    add_bold_heading(doc, '1.1 Core Network Components')

    doc.add_paragraph(
        'Core network components include enterprise-grade routers (Cisco ASR 9000 series), '
        'layer-3 switches (Juniper EX4600), and next-generation firewalls (Palo Alto PA-5200). '
        'All devices are managed via a centralized network operations center (NOC) with 24/7 '
        'monitoring by a team of 12 certified network engineers.'
    )
    doc.add_paragraph(
        'Redundant power supplies and dual uplinks are configured for all critical devices. '
        'Mean time between failures (MTBF) targets are set at 99,999 hours per chassis. '
        'Preventive maintenance schedules are enforced quarterly to minimize downtime risk.'
    )

    # Subsection 1.2 — plain bold, NOT Heading 2
    add_bold_heading(doc, '1.2 IP Addressing and VLAN Design')

    doc.add_paragraph(
        'IP addressing follows RFC 1918 private address space with a hierarchical summarization '
        'scheme. Production VLANs are allocated in the 10.10.0.0/16 range, management VLANs '
        'in the 172.16.0.0/12 range, and DMZ segments in the 192.168.0.0/24 range.'
    )
    doc.add_paragraph(
        'VLAN assignments are documented in the central IPAM database (IP Address Management). '
        'All VLAN changes require change management approval (CR) and must be logged within '
        'the ServiceNow ticketing platform with a minimum 48-hour lead time for production changes.'
    )

    doc.add_paragraph()  # blank line

    # ==========================================================================
    # SECTION 2 — plain bold, NOT Heading 1
    # ==========================================================================
    add_bold_heading(doc, 'Section 2: Security Policies and Access Control')

    doc.add_paragraph(
        'Security policy enforcement is governed by ISO 27001 and NIST SP 800-53 frameworks. '
        'All access control decisions are centrally managed through a Zero Trust Architecture '
        'model in which no user or device is implicitly trusted regardless of network position.'
    )

    # Subsection 2.1 — plain bold, NOT Heading 2
    add_bold_heading(doc, '2.1 Authentication and Authorization')

    doc.add_paragraph(
        'Multi-factor authentication (MFA) is mandatory for all privileged account access. '
        'The identity provider (IdP) is Microsoft Azure Active Directory, integrated with '
        'on-premises Active Directory via Azure AD Connect. SAML 2.0 is used for SSO federation '
        'with third-party SaaS applications including Salesforce, Workday, and ServiceNow.'
    )
    doc.add_paragraph(
        'Role-based access control (RBAC) roles are reviewed semi-annually by department heads '
        'and the Information Security team. Privilege escalation (sudo / admin) requires a '
        'separate approval workflow through the PAM (Privileged Access Management) platform, '
        'CyberArk v12.6, with full session recording enabled for all privileged sessions.'
    )

    # Subsection 2.2 — plain bold, NOT Heading 2
    add_bold_heading(doc, '2.2 Firewall and Intrusion Prevention')

    doc.add_paragraph(
        'Perimeter firewalls enforce a default-deny policy. Inbound rules are restricted to '
        'ports 443 (HTTPS) and 22 (SSH from jump hosts only). All outbound traffic is '
        'proxied through Zscaler Internet Access (ZIA) for URL filtering, SSL inspection, '
        'and DLP (Data Loss Prevention) policy enforcement.'
    )
    doc.add_paragraph(
        'Intrusion Prevention Systems (IPS) are deployed inline on all internet-facing segments. '
        'Signature updates are applied automatically every 4 hours. Threat intelligence feeds '
        'from CrowdStrike Falcon and Recorded Future are integrated to provide zero-day '
        'protection and proactive blocking of known malicious IP ranges.'
    )

    doc.add_paragraph()  # blank line

    # ==========================================================================
    # SECTION 3 — plain bold, NOT Heading 1
    # ==========================================================================
    add_bold_heading(doc, 'Section 3: Incident Response and Recovery')

    doc.add_paragraph(
        'The Incident Response (IR) lifecycle follows the NIST SP 800-61 framework: Preparation, '
        'Detection & Analysis, Containment, Eradication, Recovery, and Post-Incident Activity. '
        'All incidents are classified by severity (P1–P4) and routed to the appropriate response team.'
    )

    # Subsection 3.1 — plain bold, NOT Heading 2
    add_bold_heading(doc, '3.1 Incident Classification and Escalation')

    doc.add_paragraph(
        'P1 (Critical) incidents—those causing full service outage or confirmed data breach—require '
        'immediate notification of the CISO and CTO within 15 minutes of detection. An incident '
        'bridge call is established within 30 minutes involving representatives from Security, '
        'Network Operations, Application Support, and Executive Leadership.'
    )
    doc.add_paragraph(
        'P2 (High) incidents involve partial service degradation or suspected unauthorized access. '
        'Response SLA is 2 hours for initial containment and 8 hours for root-cause identification. '
        'P3 (Medium) and P4 (Low) incidents are managed during business hours with 24-hour and '
        '72-hour resolution targets respectively.'
    )

    # Subsection 3.2 — plain bold, NOT Heading 2
    add_bold_heading(doc, '3.2 Backup and Disaster Recovery')

    doc.add_paragraph(
        'All production data is backed up using Veeam Backup & Replication v12. Full backups run '
        'every Sunday at 02:00 UTC; incremental backups run daily at 02:00 UTC. Backup retention '
        'periods are: 30 days on-site (NetApp AFF A400 NAS), 90 days off-site (AWS S3 Glacier), '
        'and 7 years for regulatory compliance archives (Azure Blob Cold Tier).'
    )
    doc.add_paragraph(
        'Disaster recovery (DR) tests are conducted bi-annually. The RTO (Recovery Time Objective) '
        'target is 4 hours for Tier 1 applications and 24 hours for Tier 2. RPO (Recovery Point '
        'Objective) is 1 hour for Tier 1 and 4 hours for Tier 2. Results of each DR test are '
        'documented and reviewed by the Business Continuity team and executive sponsors.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
