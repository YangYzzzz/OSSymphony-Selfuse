"""
FINAL REWARD SCRIPT - SUCCESS
Task: I’m cleaning up a spec sheet in LibreOffice Writer. For Table 1, I only want the very first row—the header—to pop: bold type and text centered across each column. What’s the quickest way to do that without touching the rest of the table’s formatting?
Generated: 2025-09-10 16:11:09
Status: success
Model: azure-o3
Total Steps: 5
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

"""
Reward Script  |  LibreOffice Writer Table-Header Formatting Verification
-----------------------------------------------------------------------
Task to verify:
  The very first row (header row) of Table 1 must be:
    • Bold text in every cell
    • Text centred in every cell
  All other rows in the table must remain untouched:
    • They should NOT be bold
    • They should NOT be centred

Scoring (progressive – maximum 1.0):
    0.4 points  – Header row bold in every cell
    0.3 points  – Header row centred in every cell
    0.3 points  – All non-header rows left intact (no bold & no centring)

The script performs REAL, falsifiable checks using python-docx:
    • Iterates over every paragraph & run inside every cell
    • Confirms bold state and paragraph alignment exactly match expectations
    • Awards points only when each specific requirement is met – no natural
      conditions receive credit.
    • Returns and prints a float score, capped at 1.0, formatted
      "REWARD: X.X" as required.
"""

def _header_row_checks(table):
    """Return tuple (bold_correct, centre_correct) for the first row."""
    header = table.rows[0]
    bold_correct = True
    centre_correct = True

    for cell in header.cells:
        cell_bold_ok = True
        cell_centre_ok = True

        for paragraph in cell.paragraphs:
            # Alignment check
            if paragraph.paragraph_format.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
                cell_centre_ok = False
            # Bold check on all runs that contain visible text
            if paragraph.text.strip():
                for run in paragraph.runs:
                    if run.text.strip() and not run.bold:
                        cell_bold_ok = False

        bold_correct &= cell_bold_ok
        centre_correct &= cell_centre_ok

    return bold_correct, centre_correct


def _non_header_checks(table):
    """Ensure remaining rows are NOT bold/centred. Returns (bold_ok, centre_ok)."""
    non_bold_ok = True
    non_centre_ok = True

    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # Alignment should NOT be centred
                if paragraph.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    non_centre_ok = False
                # Runs should NOT be bold
                for run in paragraph.runs:
                    if run.text.strip() and run.bold:
                        non_bold_ok = False

    return non_bold_ok, non_centre_ok


def verify_table_header_formatting(file_path: str) -> float:
    print(f"Verifying document: {file_path}")

    if not os.path.exists(file_path):
        print("✗ File does not exist")
        return 0.0  # No score if file is missing

    try:
        doc = Document(file_path)
    except Exception as exc:
        print(f"✗ Unable to open DOCX: {exc}")
        return 0.0

    if not doc.tables:
        print("✗ No tables found in document")
        return 0.0

    # Only Table 1 is evaluated (index 0)
    table = doc.tables[0]

    score = 0.0

    # 1) Header row checks -------------------------------------------------
    header_bold_ok, header_centre_ok = _header_row_checks(table)

    if header_bold_ok:
        print("✓ Header row bold formatting correct (0.4 pts)")
        score += 0.4
    else:
        print("✗ Header row bold formatting incorrect (0 pts)")

    if header_centre_ok:
        print("✓ Header row centred correctly (0.3 pts)")
        score += 0.3
    else:
        print("✗ Header row not centred correctly (0 pts)")

    # 2) Non-header rows untouched ----------------------------------------
    non_bold_ok, non_centre_ok = _non_header_checks(table)

    if non_bold_ok and non_centre_ok:
        print("✓ Non-header rows preserved (0.3 pts)")
        score += 0.3
    else:
        if not non_bold_ok:
            print("✗ Some non-header cells are unexpectedly bold")
        if not non_centre_ok:
            print("✗ Some non-header paragraphs are centred")

    final_score = round(min(score, 1.0), 2)
    print(f"REWARD: {final_score}")
    return final_score


# -------------------------------------------------------------------------
# MAIN EXECUTION (called when script is run)                                
# -------------------------------------------------------------------------
if __name__ == "__main__":
    FILE_PATH = "/home/user/im_cleaning_up_a_spec_sheet_in_libreoffice_writer_for_table_1_i_only_want_the_very_first_rowthe_head.docx"
    verify_table_header_formatting(FILE_PATH)
