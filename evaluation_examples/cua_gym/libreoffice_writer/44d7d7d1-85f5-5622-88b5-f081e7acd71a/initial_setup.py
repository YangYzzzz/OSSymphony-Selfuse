"""
Initial Setup: Create a Writer document with a 5-column, 12-row employee data table.
Header row is dark blue (#003366) with white text. All data rows have no background color.
Task ID: writer_rd_014
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_014'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_cell_shading(cell, color_hex):
    """Set background/shading color for a table cell using XML."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shading = tc_pr.find(qn('w:shd'))
    if shading is not None:
        tc_pr.remove(shading)
    shading = tc_pr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    tc_pr.append(shading)


def set_cell_text(cell, text, bold=False, font_color=None, font_name='Calibri', font_size=11):
    """Set cell text with formatting."""
    cell.text = ''  # clear default
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if font_color:
        run.font.color.rgb = font_color


def create_initial():
    doc = Document()

    # Add a title
    heading = doc.add_heading('Employee Directory', level=1)
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add introductory paragraph
    intro = doc.add_paragraph(
        'The following table lists all current employees across departments. '
        'Please review the information and report any discrepancies to HR.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # Create table: 12 rows (1 header + 11 data) x 5 columns
    table = doc.add_table(rows=12, cols=5)
    table.style = 'Table Grid'

    # Column headers
    headers = ['Employee Name', 'Department', 'Position', 'Annual Salary', 'Start Date']

    # Header row (row 0) - dark blue background, white bold text
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        set_cell_shading(cell, '003366')
        set_cell_text(cell, header_text, bold=True,
                      font_color=RGBColor(0xFF, 0xFF, 0xFF), font_size=11)

    # Employee data (11 data rows, rows 1-11 in 0-indexed)
    employee_data = [
        ['Sarah Chen',        'Engineering',   'Senior Developer',     '$92,500',  '2021-03-15'],
        ['Marcus Johnson',    'Marketing',     'Campaign Manager',     '$78,200',  '2022-06-01'],
        ['Priya Patel',       'Finance',       'Financial Analyst',    '$85,000',  '2020-11-20'],
        ['David Kim',         'Engineering',   'DevOps Engineer',      '$95,800',  '2021-08-10'],
        ['Elena Rodriguez',   'Human Resources','HR Specialist',       '$68,500',  '2023-01-05'],
        ['James Wright',      'Sales',         'Account Executive',    '$72,000',  '2022-04-18'],
        ['Aisha Mohammed',    'Engineering',   'QA Lead',              '$88,300',  '2019-09-22'],
        ['Thomas Anderson',   'Operations',    'Logistics Coordinator','$65,400',  '2023-07-12'],
        ['Mei Lin Wang',      'Finance',       'Senior Accountant',    '$91,200',  '2020-02-28'],
        ['Carlos Gutierrez',  'Marketing',     'Content Strategist',   '$74,600',  '2022-10-03'],
        ['Rachel Foster',     'Sales',         'Regional Manager',     '$98,000',  '2018-05-14'],
    ]

    for row_idx, row_data in enumerate(employee_data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            set_cell_text(cell, value, font_size=10)

    # Add a closing paragraph
    doc.add_paragraph('')
    closing = doc.add_paragraph(
        'Last updated: March 2026. Contact HR at hr@company.com for corrections.'
    )
    closing.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    closing.runs[0].font.size = Pt(9)
    closing.runs[0].font.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
