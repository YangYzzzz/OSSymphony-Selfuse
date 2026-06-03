"""
Reward Script: Apply alternating row colors to data table
Task ID: writer_rd_014
Domain: libreoffice_writer
Scoring:
  Component 1 (0.50): Even data rows (1-based: 2,4,6,8,10,12) have light blue (#E6F0FF) background
  Component 2 (0.35): Odd data rows (1-based: 3,5,7,9,11) have white (#FFFFFF) background
  Component 3 (0.15): Header row retains dark blue (#003366) AND data rows have alternating colors applied
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_014'


def get_cell_fill(cell):
    """Extract the fill color from a cell's shading element."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    return fill.upper() if fill else None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table with 12 rows
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    if num_rows < 12:
        print(f"FAIL: Expected at least 12 rows, found {num_rows}")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Table has {num_rows} rows and {num_cols} columns")

    # Component 1: Even data rows have light blue (#E6F0FF) background (0.50 points)
    # Even data rows in 1-based numbering: 2, 4, 6, 8, 10, 12 → 0-based indices: 1, 3, 5, 7, 9, 11
    try:
        even_row_indices = [1, 3, 5, 7, 9, 11]
        even_pass_count = 0
        even_total_cells = 0

        for row_idx in even_row_indices:
            if row_idx >= num_rows:
                continue
            row = table.rows[row_idx]
            for cell in row.cells:
                fill = get_cell_fill(cell)
                even_total_cells += 1
                if fill == 'E6F0FF':
                    even_pass_count += 1

        if even_total_cells > 0 and even_pass_count == even_total_cells:
            print(f"PASS: Component 1 — All {even_pass_count} cells in even data rows have #E6F0FF (0.50 pts)")
            total_score += 0.50
        elif even_total_cells > 0 and even_pass_count > 0:
            partial = 0.50 * (even_pass_count / even_total_cells)  # proportional credit
            if partial > 0:
                print(f"PARTIAL: Component 1 — {even_pass_count}/{even_total_cells} cells have #E6F0FF ({partial:.2f} pts)")
                total_score += partial
        else:
            print(f"FAIL: Component 1 — No cells in even data rows have #E6F0FF (0/{even_total_cells})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Odd data rows have white (#FFFFFF) background (0.35 points)
    # Odd data rows in 1-based: 3, 5, 7, 9, 11 → 0-based indices: 2, 4, 6, 8, 10
    try:
        odd_row_indices = [2, 4, 6, 8, 10]
        odd_pass_count = 0
        odd_total_cells = 0

        for row_idx in odd_row_indices:
            if row_idx >= num_rows:
                continue
            row = table.rows[row_idx]
            for cell in row.cells:
                fill = get_cell_fill(cell)
                odd_total_cells += 1
                if fill == 'FFFFFF':
                    odd_pass_count += 1

        if odd_total_cells > 0 and odd_pass_count == odd_total_cells:
            print(f"PASS: Component 2 — All {odd_pass_count} cells in odd data rows have #FFFFFF (0.35 pts)")
            total_score += 0.35
        elif odd_total_cells > 0 and odd_pass_count > 0:
            partial = 0.35 * (odd_pass_count / odd_total_cells)  # proportional credit
            if partial > 0:
                print(f"PARTIAL: Component 2 — {odd_pass_count}/{odd_total_cells} cells have #FFFFFF ({partial:.2f} pts)")
                total_score += partial
        else:
            print(f"FAIL: Component 2 — No cells in odd data rows have #FFFFFF (0/{odd_total_cells})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header row retains dark blue (#003366) AND alternating colors are applied (0.15 points)
    # This is a compound check: header preservation is only scored when data rows also changed.
    # The "data rows changed" part ensures this fails on the initial env.
    try:
        header_row = table.rows[0]
        header_fills = [get_cell_fill(cell) for cell in header_row.cells]
        header_correct = all(f == '003366' for f in header_fills)
        if not header_correct:
            print(f"FAIL: Component 3 — Header cell fills are {header_fills}, expected all 003366")

        # Check that at least some data rows have shading applied (not None)
        data_row_fills = [get_cell_fill(table.rows[r].cells[0]) for r in range(1, min(num_rows, 12))]
        data_rows_have_shading = any(f is not None for f in data_row_fills)

        if header_correct and data_rows_have_shading:
            print(f"PASS: Component 3 — Header retains #003366 and data rows have alternating colors (0.15 pts)")
            total_score += 0.15
        elif not header_correct:
            print(f"FAIL: Component 3 — Header row background is not #003366")
        else:
            print(f"FAIL: Component 3 — No alternating colors applied to data rows")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice changes before verification
def persist_app_state(domain):
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
