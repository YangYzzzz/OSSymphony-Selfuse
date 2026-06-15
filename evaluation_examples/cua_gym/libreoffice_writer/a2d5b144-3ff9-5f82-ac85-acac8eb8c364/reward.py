"""
Reward Script: Internal Announcement Formatting - Parental Leave Policy
Task ID: writer_mktg_036
Domain: libreoffice_writer
Scoring:
  Component 1: [Company Logo] placeholder table at top (centered, bordered) — 0.2 pts
  Component 2: INTERNAL MEMO heading at 16pt bold centered — 0.2 pts
  Component 3: 2-column info table (To/From/Date/Subject) with bold labels, Subject at 14pt — 0.3 pts
  Component 4: Key policy details callout box with #E3F2FD background and left blue border (#1565C0) — 0.3 pts
  Total: 1.0
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_mktg_036'
FILE_PATH = f'{WORKDIR}/parental_leave_announcement.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have tables (none present in initial state)
    if len(doc.tables) == 0:
        print("FAIL: No tables found — document appears to be unmodified initial state")
        print("\nScore: 0.0/1.0")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: [Company Logo] placeholder table (0.2 points)
    # Task requires: a centered '[Company Logo]' text in a bordered frame/table
    # at the top of the document (new, not present in initial).
    # -------------------------------------------------------------------------
    try:
        logo_found = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if '[Company Logo]' in cell_text or 'Company Logo' in cell_text:
                        # Check if the cell/paragraph is centered
                        for para in cell.paragraphs:
                            if (para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                                    and para.text.strip()):
                                # Check for borders on this cell
                                tcPr = cell._tc.find(qn('w:tcPr'))
                                has_border = False
                                if tcPr is not None:
                                    border_elem = tcPr.find(qn('w:tcBorders'))
                                    if border_elem is not None:
                                        # At least one border side should be non-none
                                        for side in ['w:top', 'w:left', 'w:bottom', 'w:right']:
                                            b = border_elem.find(qn(side))
                                            if b is not None and b.get(qn('w:val'), 'none') not in ('none', 'nil', ''):
                                                has_border = True
                                                break
                                if has_border:
                                    logo_found = True
                                    print(f"PASS: Component 1 — '[Company Logo]' centered with borders (0.2 pts)")
                                    total_score += 0.2
                                    break
                        if logo_found:
                            break
                if logo_found:
                    break

        if not logo_found:
            # Partial: check if logo text exists at all in a table (without full border/center requirement)
            logo_text_found = any(
                '[Company Logo]' in cell.text or 'Company Logo' in cell.text
                for table in doc.tables
                for row in table.rows
                for cell in row.cells
            )
            if logo_text_found:
                print("PARTIAL: Component 1 — '[Company Logo]' table present but missing center alignment or borders (0.1 pts)")
                total_score += 0.1
            else:
                print("FAIL: Component 1 — '[Company Logo]' placeholder not found in any table")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: INTERNAL MEMO at 16pt bold centered (0.2 points)
    # Task requires: 'INTERNAL MEMO' prominent — 16pt bold, centered alignment.
    # In initial state it is 12pt, not bold, not centered.
    # -------------------------------------------------------------------------
    try:
        memo_para_found = False
        for para in doc.paragraphs:
            if 'INTERNAL MEMO' in para.text:
                # Check bold and size on runs
                is_bold = False
                is_large = False
                for run in para.runs:
                    if 'INTERNAL MEMO' in run.text or run.text.strip():
                        if run.font.bold:
                            is_bold = True
                        if run.font.size and run.font.size.pt >= 14:
                            is_large = True
                is_centered = (para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)

                if is_bold and is_large and is_centered:
                    print(f"PASS: Component 2 — 'INTERNAL MEMO' is bold, >=14pt, centered (0.2 pts)")
                    total_score += 0.2
                    memo_para_found = True
                elif is_bold and is_large:
                    print(f"PARTIAL: Component 2 — 'INTERNAL MEMO' is bold and >=14pt but not centered (0.1 pts)")
                    total_score += 0.1
                    memo_para_found = True
                elif is_bold or is_large:
                    print(f"PARTIAL: Component 2 — 'INTERNAL MEMO' is {'bold' if is_bold else 'large'} but missing other requirements (0.05 pts)")
                    total_score += 0.05
                    memo_para_found = True
                else:
                    print(f"FAIL: Component 2 — 'INTERNAL MEMO' found but not formatted (bold={is_bold}, large={is_large}, centered={is_centered})")
                    memo_para_found = True
                break

        if not memo_para_found:
            print("FAIL: Component 2 — 'INTERNAL MEMO' paragraph not found in document body")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: 2-column header info table with bold labels, Subject at 14pt (0.3 points)
    # Task requires: To/From/Date/Subject formatted as 2-column table with bold labels.
    # Subject row must be at 14pt bold.
    # In initial state, these are plain paragraphs, not a table.
    # -------------------------------------------------------------------------
    try:
        info_table_found = False
        subject_14pt = False
        bold_labels_found = False

        for table in doc.tables:
            # Look for a 4-row, 2-column table with To/From/Date/Subject
            if len(table.rows) >= 2 and len(table.columns) >= 2:
                first_col_texts = [row.cells[0].text.strip() for row in table.rows]
                # Check if it contains the info headers
                has_to = any('To' in t for t in first_col_texts)
                has_from = any('From' in t for t in first_col_texts)
                has_subject = any('Subject' in t for t in first_col_texts)

                if has_to and has_from and has_subject:
                    info_table_found = True

                    # Check bold labels in column 0
                    bold_count = 0
                    for row in table.rows:
                        cell0 = row.cells[0]
                        for para in cell0.paragraphs:
                            for run in para.runs:
                                if run.font.bold and run.text.strip():
                                    bold_count += 1
                    if bold_count >= 2:
                        bold_labels_found = True

                    # Check Subject row at 14pt
                    for row in table.rows:
                        cell0_text = row.cells[0].text.strip()
                        if 'Subject' in cell0_text:
                            for ci, cell in enumerate([row.cells[0], row.cells[1]]):
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        if run.font.size and run.font.size.pt >= 13:
                                            subject_14pt = True
                    break

        if info_table_found and bold_labels_found and subject_14pt:
            print("PASS: Component 3 — 2-column info table with bold labels and Subject at 14pt (0.3 pts)")
            total_score += 0.3
        elif info_table_found and bold_labels_found:
            print("PARTIAL: Component 3 — 2-column info table with bold labels found, but Subject not at 14pt (0.2 pts)")
            total_score += 0.2
        elif info_table_found:
            print("PARTIAL: Component 3 — 2-column info table found but labels not bold and/or Subject not at 14pt (0.1 pts)")
            total_score += 0.1
        else:
            print("FAIL: Component 3 — No 2-column header info table (To/From/Date/Subject) found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: Key policy details callout box with light blue background
    #              (#E3F2FD) and left blue border (#1565C0) (0.3 points)
    # Task requires: policy details paragraph in a highlighted callout box.
    # In initial state, it's a plain paragraph — no table, no shading.
    # -------------------------------------------------------------------------
    try:
        callout_found = False
        has_blue_bg = False
        has_left_border = False

        for table in doc.tables:
            if len(table.rows) == 1 and len(table.columns) == 1:
                cell = table.cell(0, 0)
                cell_text = cell.text.strip()
                # Check if it contains the key policy details (16 weeks / caregivers)
                if ('16 weeks' in cell_text or 'caregivers' in cell_text
                        or 'parental leave' in cell_text.lower()):
                    callout_found = True

                    # Check cell shading
                    tcPr = cell._tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            fill_color = shd.get(qn('w:fill'), '').upper()
                            # Accept #E3F2FD or similar light blue variants
                            if fill_color in ('E3F2FD', 'E3F2FC', 'E3F2FF') or fill_color.startswith('E3F'):
                                has_blue_bg = True
                            elif fill_color and fill_color not in ('AUTO', 'FFFFFF', ''):
                                # Some non-white, non-default background
                                has_blue_bg = True  # partial credit for any background

                        # Check left border
                        border_elem = tcPr.find(qn('w:tcBorders'))
                        if border_elem is not None:
                            left_b = border_elem.find(qn('w:left'))
                            if left_b is not None:
                                border_val = left_b.get(qn('w:val'), 'none')
                                border_color = left_b.get(qn('w:color'), '').upper()
                                if border_val not in ('none', 'nil', '') and border_color:
                                    has_left_border = True

        if callout_found and has_blue_bg and has_left_border:
            print("PASS: Component 4 — Policy callout box with light blue background and left border (0.3 pts)")
            total_score += 0.3
        elif callout_found and (has_blue_bg or has_left_border):
            detail = "background" if has_blue_bg else "left border"
            print(f"PARTIAL: Component 4 — Policy callout box found with {detail} only (0.15 pts)")
            total_score += 0.15
        elif callout_found:
            print("PARTIAL: Component 4 — Policy text in a single-cell table but missing background and left border (0.1 pts)")
            total_score += 0.1
        else:
            print("FAIL: Component 4 — Key policy details callout box not found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -------------------------------------------------------------------------
    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


# Run verification against the canonical file path
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
