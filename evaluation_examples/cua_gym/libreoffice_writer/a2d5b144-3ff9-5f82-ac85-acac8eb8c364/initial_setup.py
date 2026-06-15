"""
Initial Setup: Internal Memo - Parental Leave Announcement (plain, unformatted)
Task ID: writer_mktg_036
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'  # VM path — file goes on Desktop per task spec
TASK_ID = 'writer_mktg_036'
OUTPUT = f'{WORKDIR}/parental_leave_announcement.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set default font to 12pt throughout
    style = doc.styles['Normal']
    style.font.size = Pt(12)

    # --- Line 1: INTERNAL MEMO (plain 12pt, no special formatting) ---
    p_memo = doc.add_paragraph()
    run_memo = p_memo.add_run('INTERNAL MEMO')
    run_memo.font.size = Pt(12)

    # --- Header fields (plain text, no table) ---
    p_to = doc.add_paragraph()
    run_to = p_to.add_run('To: All Employees')
    run_to.font.size = Pt(12)

    p_from = doc.add_paragraph()
    run_from = p_from.add_run('From: HR & Communications')
    run_from.font.size = Pt(12)

    p_date = doc.add_paragraph()
    run_date = p_date.add_run('Date: March 4, 2026')
    run_date.font.size = Pt(12)

    p_subject = doc.add_paragraph()
    run_subject = p_subject.add_run('Subject: Updated Parental Leave Policy')
    run_subject.font.size = Pt(12)

    # --- 4 Body Paragraphs ---
    # Paragraph 1: Introduction
    p1 = doc.add_paragraph()
    run1 = p1.add_run(
        'We are pleased to announce an important update to our parental leave policy, '
        'effective April 1, 2026. After careful review and listening to employee feedback, '
        'we have made significant enhancements to better support our team members during '
        'one of life\'s most meaningful transitions.'
    )
    run1.font.size = Pt(12)

    # Paragraph 2: Key policy details (THIS IS THE CALLOUT CONTENT)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        '16 weeks paid leave for primary caregivers, 8 weeks for secondary caregivers, '
        'effective April 1, 2026.'
    )
    run2.font.size = Pt(12)

    # Paragraph 3: Eligibility
    p3 = doc.add_paragraph()
    run3 = p3.add_run(
        'All full-time employees who have completed at least six months of continuous service '
        'are eligible for this updated policy. This applies to birth parents, adoptive parents, '
        'and foster parents alike. Part-time employees working 20 or more hours per week are '
        'eligible for pro-rated leave benefits.'
    )
    run3.font.size = Pt(12)

    # Paragraph 4: Next steps / closing
    p4 = doc.add_paragraph()
    run4 = p4.add_run(
        'Please reach out to your HR Business Partner or visit our internal HR portal for '
        'full policy details, FAQs, and the leave request form. We encourage all employees '
        'to review the complete policy documentation. If you have questions or need support, '
        'the HR team is available at hr@company.com or ext. 4200.'
    )
    run4.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
