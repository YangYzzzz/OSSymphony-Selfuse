"""
Initial Setup: Format pricing table alignment and add Total row
Task ID: writer_biz_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_031'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("Service Proposal — Q2 2025", level=1)

    # Intro paragraph
    doc.add_paragraph(
        "Dear Ms. Nakamura,\n\n"
        "Thank you for the opportunity to present our consulting proposal for "
        "the digital transformation initiative at Greenfield Industries. Below "
        "is a detailed breakdown of the services and associated costs."
    )

    doc.add_paragraph("")  # spacer

    # Pricing table: 2 columns, header + 5 data rows = 6 rows total
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Description"
    hdr_cells[1].text = "Amount"

    # Data rows — all left-aligned (default), realistic content
    data = [
        ("Discovery & Requirements Analysis", "$4,500.00"),
        ("System Architecture Design", "$8,750.00"),
        ("Custom Software Development", "$22,300.00"),
        ("Quality Assurance & Testing", "$6,200.00"),
        ("Deployment & Training", "$3,250.00"),
    ]

    for i, (desc, amt) in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = desc
        row_cells[1].text = amt

    # Ensure ALL cells are explicitly left-aligned (the pre-task state)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Closing paragraph
    doc.add_paragraph("")
    doc.add_paragraph(
        "We look forward to partnering with Greenfield Industries on this "
        "exciting project. Please do not hesitate to reach out with any questions."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph("Elena Vasquez")
    doc.add_paragraph("Senior Consultant, Apex Digital Solutions")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
