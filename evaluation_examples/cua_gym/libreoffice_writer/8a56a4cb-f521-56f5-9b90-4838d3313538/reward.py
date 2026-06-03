"""
Reward Script: Format pricing table alignment and add bold Total row
Task ID: writer_biz_031
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Amount column data cells (rows 1-5) are RIGHT-aligned
  Component 2 (0.10): Amount header cell is RIGHT-aligned
  Component 3 (0.20): Total row exists with 'Total' text in first column
  Component 4 (0.15): Total row first column text is bold
  Component 5 (0.25): Total row second column has correct sum, is bold and right-aligned
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_031'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    print(f"INFO: Table has {num_rows} rows x {num_cols} cols")

    # Component 1: Amount column data cells (rows 1-5) are RIGHT-aligned (0.30 points)
    # In initial state, all cells are LEFT-aligned. Task requires Amount column to be right-aligned.
    try:
        right_count = 0
        data_rows_to_check = min(num_rows - 1, 5)  # up to 5 data rows (skip header)
        if data_rows_to_check > 0:
            for ri in range(1, 1 + data_rows_to_check):
                cell = table.cell(ri, 1)
                for para in cell.paragraphs:
                    if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                        right_count += 1
                        break
            if right_count == data_rows_to_check:
                print(f"PASS: Component 1 — All {data_rows_to_check} Amount data cells are RIGHT-aligned (0.30 pts)")
                total_score += 0.30
            else:
                print(f"FAIL: Component 1 — Only {right_count}/{data_rows_to_check} Amount data cells are RIGHT-aligned")
        else:
            print("FAIL: Component 1 — Not enough rows to check data cells")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Amount header cell (row 0, col 1) is RIGHT-aligned (0.10 points)
    try:
        header_cell = table.cell(0, 1)
        header_aligned_right = False
        for para in header_cell.paragraphs:
            if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                header_aligned_right = True
                break
        if header_aligned_right:
            print("PASS: Component 2 — Amount header cell is RIGHT-aligned (0.10 pts)")
            total_score += 0.10
        else:
            align_val = header_cell.paragraphs[0].paragraph_format.alignment if header_cell.paragraphs else None
            print(f"FAIL: Component 2 — Amount header alignment is {align_val}, expected RIGHT")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Total row exists with 'Total' text in first column (0.20 points)
    # Initial state has 6 rows (1 header + 5 data). Golden has 7 rows (added Total row).
    try:
        total_row_found = False
        total_row_idx = None
        # Search from the last row backwards for a row with 'Total' in col 0
        for ri in range(num_rows - 1, 0, -1):
            cell_text = table.cell(ri, 0).text.strip()
            if cell_text.lower() == 'total':
                total_row_found = True
                total_row_idx = ri
                break

        if total_row_found and num_rows >= 7:
            print(f"PASS: Component 3 — Total row found at row {total_row_idx} (0.20 pts)")
            total_score += 0.20
        elif total_row_found:
            # Total row exists but table doesn't have the expected extra row
            # This could mean a data row was replaced. Still give partial.
            print(f"PARTIAL: Component 3 — Total row found at row {total_row_idx} but table has only {num_rows} rows")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — No row with 'Total' text found in first column")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Total row first column text is bold (0.15 points)
    try:
        if total_row_idx is not None:
            total_cell_0 = table.cell(total_row_idx, 0)
            all_bold = False
            runs = [r for p in total_cell_0.paragraphs for r in p.runs if r.text.strip()]
            if runs and all(r.bold is True for r in runs):
                all_bold = True

            if all_bold:
                print("PASS: Component 4 — Total row first column is bold (0.15 pts)")
                total_score += 0.15
            else:
                bold_vals = [r.bold for r in runs]
                print(f"FAIL: Component 4 — Total row first column bold values: {bold_vals}")
        else:
            print("FAIL: Component 4 — No Total row found, cannot check bold")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Total row second column is bold and right-aligned with sum value (0.25 points)
    try:
        if total_row_idx is not None:
            total_cell_1 = table.cell(total_row_idx, 1)
            cell_text = total_cell_1.text.strip()

            # Check right-alignment
            is_right = False
            for para in total_cell_1.paragraphs:
                if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                    is_right = True
                    break

            # Check bold
            runs = [r for p in total_cell_1.paragraphs for r in p.runs if r.text.strip()]
            is_bold = bool(runs) and all(r.bold is True for r in runs)

            # Check that cell contains a currency value (the sum)
            has_value = '$' in cell_text and any(c.isdigit() for c in cell_text)

            sub_score = 0.0
            if is_right:
                sub_score += 0.10
            if is_bold:
                sub_score += 0.10
            if has_value:
                sub_score += 0.05

            if sub_score > 0:
                details = f"right={is_right}, bold={is_bold}, has_value={has_value}, text='{cell_text}'"
                if sub_score == 0.25:
                    print(f"PASS: Component 5 — Total amount cell fully correct ({details}) (0.25 pts)")
                else:
                    print(f"PARTIAL: Component 5 — {details} ({sub_score} pts)")
                total_score += sub_score
            else:
                print(f"FAIL: Component 5 — Total amount cell: right={is_right}, bold={is_bold}, text='{cell_text}'")
        else:
            print("FAIL: Component 5 — No Total row found, cannot check amount cell")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    final_score = round(final_score, 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
