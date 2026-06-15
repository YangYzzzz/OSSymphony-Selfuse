"""
Initial Setup: Master document with 3 chapter sections, no headers configured.
Task ID: writer_rm_089
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_089'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup defaults ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # =============================================
    # SECTION 1 / CHAPTER 1: Introduction
    # =============================================
    # Title
    h1 = doc.add_heading('Chapter 1: Introduction', level=1)
    h1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph(
        'The purpose of this thesis is to investigate the effects of urbanization '
        'on local biodiversity in temperate regions. Over the past three decades, '
        'rapid urban expansion has transformed natural landscapes, leading to habitat '
        'fragmentation and species displacement across multiple taxonomic groups.'
    )
    doc.add_paragraph(
        'Previous research by Martinez et al. (2019) demonstrated that urban green '
        'corridors can mitigate some of these effects, though the extent of recovery '
        'remains poorly quantified. This study builds upon their framework by '
        'incorporating longitudinal data from twelve metropolitan areas.'
    )
    doc.add_paragraph(
        'The central hypothesis posits that cities implementing integrated green '
        'infrastructure policies exhibit measurably higher species richness indices '
        'compared to those following conventional development patterns. We test this '
        'across avian, lepidopteran, and small mammal populations.'
    )
    doc.add_paragraph(
        'This chapter provides an overview of the research questions, outlines '
        'the scope of the investigation, and describes the structure of subsequent '
        'chapters. Section 1.1 presents the problem statement, Section 1.2 discusses '
        'the significance of this research, and Section 1.3 defines key terminology.'
    )

    # =============================================
    # SECTION 2 / CHAPTER 2: Background
    # =============================================
    # Add a new section with a page break
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    h2 = doc.add_heading('Chapter 2: Background', level=1)
    h2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph(
        'Urban ecology has emerged as a critical subdiscipline within conservation '
        'biology since the early 2000s. The foundational work of Pickett et al. (2001) '
        'established the Baltimore Ecosystem Study, which pioneered the integration '
        'of social and ecological research in urban settings.'
    )
    doc.add_paragraph(
        'Subsequent studies expanded this framework to include cities across different '
        'climatic zones. Notable contributions include the Stockholm Resilience Centre\'s '
        'analysis of urban tipping points (Folke et al., 2010) and the Shanghai Urban '
        'Biodiversity Assessment conducted between 2015 and 2020.'
    )
    doc.add_paragraph(
        'The concept of ecosystem services in urban environments gained traction '
        'following the Millennium Ecosystem Assessment (2005). Cities were recognized '
        'not merely as consumers of ecosystem services but as complex systems capable '
        'of generating and maintaining ecological functions through careful planning.'
    )
    doc.add_paragraph(
        'Green infrastructure, defined by the European Commission as "a strategically '
        'planned network of natural and semi-natural areas," has become a central policy '
        'tool. Studies by Benedict and McMahon (2006) and more recently by Pauleit et al. '
        '(2021) demonstrate measurable improvements in urban biodiversity metrics where '
        'such infrastructure is implemented at scale.'
    )
    doc.add_paragraph(
        'Despite these advances, significant gaps remain in our understanding of how '
        'species respond to urban gradients over extended time periods. Most existing '
        'studies are cross-sectional, capturing a single point in time rather than '
        'tracking population dynamics across decades of urban development.'
    )

    # =============================================
    # SECTION 3 / CHAPTER 3: Methodology
    # =============================================
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    h3 = doc.add_heading('Chapter 3: Methodology', level=1)
    h3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph(
        'This study employs a mixed-methods approach combining quantitative field '
        'surveys with geospatial analysis. Data collection occurred across twelve '
        'metropolitan areas in North America and Western Europe between 2018 and 2024.'
    )
    doc.add_paragraph(
        'Bird populations were surveyed using standardized point-count methodology '
        'following protocols established by the North American Breeding Bird Survey. '
        'Each survey site consisted of a 500-meter radius circle with observation '
        'periods of exactly 10 minutes per point, conducted between 0600 and 1000 hours.'
    )
    doc.add_paragraph(
        'Butterfly surveys followed the Pollard Walk transect method, with weekly '
        'counts along predetermined 2-kilometer routes during the active season '
        '(April through September in the Northern Hemisphere). Species identification '
        'was confirmed through photographic documentation and expert review.'
    )
    doc.add_paragraph(
        'Small mammal populations were assessed using Sherman live-trap grids deployed '
        'in a 7x7 configuration with 15-meter spacing. Trapping sessions lasted three '
        'consecutive nights per site per season, with individuals marked using PIT tags '
        'to enable capture-recapture analysis.'
    )
    doc.add_paragraph(
        'Geospatial analysis utilized Landsat 8 satellite imagery processed through '
        'Google Earth Engine to calculate the Normalized Difference Vegetation Index '
        '(NDVI) and impervious surface percentages for each study site. Land use '
        'classification followed the CORINE methodology adapted for urban environments.'
    )
    doc.add_paragraph(
        'Statistical analysis was performed using R version 4.3.1 with packages '
        'including lme4 for mixed-effects models, vegan for community ecology metrics, '
        'and sf/terra for spatial operations. All models controlled for latitude, '
        'elevation, and regional climate variables derived from WorldClim 2.1 datasets.'
    )

    # Ensure NO headers are set on any section
    for section in doc.sections:
        section.header.is_linked_to_previous = True
        # Clear any default header content
        for para in section.header.paragraphs:
            para.text = ''
        section.footer.is_linked_to_previous = True
        for para in section.footer.paragraphs:
            para.text = ''

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
