"""
Reward Script: HR Metrics & Analytics Dashboard Report
Task ID: writer_hr_095
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Document has 8+ tables (initial has 0)
  Component 2 (0.15): Table of Contents section with 8 entries
  Component 3 (0.20): KPI table has 12 data rows with 5 columns including trend indicators
  Component 4 (0.15): Department table has 10 departments
  Component 5 (0.10): Monthly turnover table has 12 months with voluntary/involuntary breakdown
  Component 6 (0.10): All tables have captions (paragraphs starting with "Table N:")
  Component 7 (0.10): Bold header rows in tables (executive styling)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_095'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_tables = len(doc.tables)

    # Component 1: Document has 8+ tables (0.20 points)
    # Initial has 0 tables; golden has 8. This is the core transformation.
    try:
        if num_tables >= 8:
            print(f"PASS: Component 1 — Document has {num_tables} tables (>= 8) (0.20 pts)")
            total_score += 0.20
        elif num_tables >= 5:
            partial = 0.10
            print(f"PARTIAL: Component 1 — Document has {num_tables} tables (>= 5 but < 8) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — Document has {num_tables} tables, expected >= 8")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table of Contents section with 8 entries (0.15 points)
    # Initial has no TOC; golden has a Heading 1 "Table of Contents" followed by 8 entries.
    try:
        toc_heading_found = False
        toc_entries = []
        for i, p in enumerate(doc.paragraphs):
            if p.style and p.style.name == 'Heading 1' and 'table of contents' in p.text.lower():
                toc_heading_found = True
                # Collect subsequent entries until next heading
                j = i + 1
                while j < len(doc.paragraphs) and (doc.paragraphs[j].style is None or doc.paragraphs[j].style.name != 'Heading 1'):
                    text = doc.paragraphs[j].text.strip()
                    if text and text.lower().startswith('table'):
                        toc_entries.append(text)
                    j += 1
                break

        if toc_heading_found and len(toc_entries) >= 8:
            print(f"PASS: Component 2 — TOC heading found with {len(toc_entries)} entries (0.15 pts)")
            total_score += 0.15
        elif toc_heading_found and len(toc_entries) >= 4:
            partial = 0.08
            print(f"PARTIAL: Component 2 — TOC heading found with {len(toc_entries)} entries ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — TOC heading: {toc_heading_found}, entries: {len(toc_entries)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: KPI table has 12 data rows with 5 columns and trend indicators (0.20 points)
    # Initial has KPI data in paragraphs; golden has it in a 13-row x 5-col table (1 header + 12 metrics)
    try:
        if num_tables >= 1:
            kpi_table = doc.tables[0]
            kpi_rows = len(kpi_table.rows)
            kpi_cols = len(kpi_table.columns)
            # Check header matches expected structure
            header_cells = [c.text.strip().lower() for c in kpi_table.rows[0].cells]
            has_metric_col = any('metric' in h for h in header_cells)
            has_trend_col = any('trend' in h for h in header_cells)

            # Check trend indicators exist in data rows (arrows or up/down text)
            trend_count = 0
            for row in kpi_table.rows[1:]:
                last_cell = row.cells[-1].text.strip().lower()
                if 'up' in last_cell or 'down' in last_cell or '↑' in last_cell or '↓' in last_cell:
                    trend_count += 1

            data_rows = kpi_rows - 1  # subtract header
            score_3 = 0.0
            if data_rows >= 12 and kpi_cols >= 5 and has_metric_col and has_trend_col:
                score_3 = 0.12
            if trend_count >= 12:
                score_3 += 0.08

            if score_3 > 0:
                print(f"PASS: Component 3 — KPI table: {data_rows} data rows, {kpi_cols} cols, {trend_count} trend indicators ({score_3} pts)")
                total_score += score_3
            else:
                print(f"FAIL: Component 3 — KPI table: {data_rows} data rows, {kpi_cols} cols, trends={trend_count}, metric_col={has_metric_col}, trend_col={has_trend_col}")
        else:
            print(f"FAIL: Component 3 — No tables found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Department table has 10 departments (0.15 points)
    # Initial has department data in paragraphs; golden has table with 10 dept rows
    try:
        dept_found = False
        if num_tables >= 2:
            dept_table = doc.tables[1]
            dept_header = [c.text.strip().lower() for c in dept_table.rows[0].cells]
            has_dept_col = any('department' in h for h in dept_header)
            has_headcount_col = any('headcount' in h for h in dept_header)

            if has_dept_col and has_headcount_col:
                # Count actual department rows (exclude header and total rows)
                dept_count = 0
                for row in dept_table.rows[1:]:
                    first_cell = row.cells[0].text.strip().lower()
                    if first_cell and first_cell != 'total':
                        dept_count += 1

                if dept_count >= 10:
                    print(f"PASS: Component 4 — Department table has {dept_count} departments (0.15 pts)")
                    total_score += 0.15
                    dept_found = True
                elif dept_count >= 5:
                    partial = 0.08
                    print(f"PARTIAL: Component 4 — Department table has {dept_count} departments ({partial} pts)")
                    total_score += partial
                    dept_found = True

        if not dept_found and num_tables >= 2:
            # Try searching all tables for one with department structure
            for ti, table in enumerate(doc.tables):
                header = [c.text.strip().lower() for c in table.rows[0].cells]
                if any('department' in h for h in header) and any('headcount' in h for h in header):
                    dept_count = sum(1 for row in table.rows[1:] if row.cells[0].text.strip().lower() not in ('', 'total'))
                    if dept_count >= 10:
                        print(f"PASS: Component 4 — Department table (table {ti}) has {dept_count} departments (0.15 pts)")
                        total_score += 0.15
                        dept_found = True
                        break
        if not dept_found:
            print(f"FAIL: Component 4 — No department table with 10+ departments found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Monthly turnover table with 12 months and voluntary/involuntary (0.10 points)
    # Initial has turnover data in paragraphs; golden has structured table
    try:
        turnover_found = False
        if num_tables >= 5:
            for ti, table in enumerate(doc.tables):
                header = [c.text.strip().lower() for c in table.rows[0].cells]
                has_month = any('month' in h for h in header)
                has_voluntary = any('voluntary' in h for h in header)
                has_involuntary = any('involuntary' in h for h in header)

                if has_month and has_voluntary and has_involuntary:
                    # Count month rows (exclude header and summary rows)
                    month_names = {'january', 'february', 'march', 'april', 'may', 'june',
                                   'july', 'august', 'september', 'october', 'november', 'december'}
                    month_count = 0
                    for row in table.rows[1:]:
                        first_cell = row.cells[0].text.strip().lower()
                        if first_cell in month_names:
                            month_count += 1

                    if month_count >= 12:
                        print(f"PASS: Component 5 — Turnover table (table {ti}) has {month_count} months with vol/invol breakdown (0.10 pts)")
                        total_score += 0.10
                        turnover_found = True
                        break

        if not turnover_found:
            print(f"FAIL: Component 5 — No monthly turnover table with 12 months and voluntary/involuntary breakdown found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: All tables have captions (0.10 points)
    # Initial has no tables so no captions. Golden has "Table N: ..." paragraphs before each table.
    try:
        import re
        caption_pattern = re.compile(r'^Table\s+\d+\s*:', re.IGNORECASE)
        caption_paras = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if caption_pattern.match(text):
                caption_paras.append(text)

        # Deduplicate (TOC and section captions may repeat)
        # We need at least 8 unique caption texts corresponding to the 8 tables
        unique_captions = set()
        for c in caption_paras:
            # Normalize
            unique_captions.add(c.strip())

        # Each table number should appear in captions (outside the TOC)
        # Look for captions that appear AFTER the TOC section (near actual tables)
        # Simple check: at least 8 caption paragraphs exist
        if len(caption_paras) >= 8 and num_tables >= 8:
            print(f"PASS: Component 6 — Found {len(caption_paras)} table caption paragraphs for {num_tables} tables (0.10 pts)")
            total_score += 0.10
        elif len(caption_paras) >= 4:
            partial = 0.05
            print(f"PARTIAL: Component 6 — Found {len(caption_paras)} table captions ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 6 — Found {len(caption_paras)} table captions, expected >= 8")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Bold header rows in tables (executive styling) (0.10 points)
    # Initial has no tables. Golden has bold header rows in all tables.
    try:
        if num_tables >= 8:
            tables_with_bold_headers = 0
            for table in doc.tables:
                header_row = table.rows[0]
                has_any_bold = False
                for cell in header_row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.bold:
                                has_any_bold = True
                                break
                        if has_any_bold:
                            break
                    if has_any_bold:
                        break
                if has_any_bold:
                    tables_with_bold_headers += 1

            if tables_with_bold_headers >= 8:
                print(f"PASS: Component 7 — {tables_with_bold_headers}/{num_tables} tables have bold header rows (0.10 pts)")
                total_score += 0.10
            elif tables_with_bold_headers >= 4:
                partial = 0.05
                print(f"PARTIAL: Component 7 — {tables_with_bold_headers}/{num_tables} tables have bold header rows ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 7 — Only {tables_with_bold_headers}/{num_tables} tables have bold header rows")
        else:
            print(f"FAIL: Component 7 — Not enough tables ({num_tables}) to check headers")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
