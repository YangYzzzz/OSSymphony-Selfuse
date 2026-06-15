"""
Initial Setup: HR Analytics Dashboard Report - Raw data in paragraph form
Task ID: writer_hr_095
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_095'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('HR_Analytics_Dashboard_Q4_2025', level=0)

    doc.add_paragraph(
        'This document contains raw HR metrics and analytics data for the fourth quarter of 2025. '
        'The data below needs to be organized into professional tables with executive-appropriate styling, '
        'captions, and a master table of contents.'
    )

    # --- Section: KPI Summary ---
    doc.add_heading('KPI Summary Data', level=1)
    doc.add_paragraph(
        'The following HR key performance indicators were tracked during Q4 2025:'
    )
    kpi_data = [
        ("Employee Turnover Rate", "4.2%", "3.5%", "+0.7%", "Up"),
        ("Time to Fill (days)", "38", "30", "+8", "Up"),
        ("Offer Acceptance Rate", "87%", "90%", "-3%", "Down"),
        ("Employee Satisfaction Score", "4.1/5.0", "4.3/5.0", "-0.2", "Down"),
        ("Training Hours per Employee", "24", "20", "+4", "Up"),
        ("Absenteeism Rate", "2.8%", "2.5%", "+0.3%", "Up"),
        ("Revenue per Employee ($K)", "285", "300", "-15", "Down"),
        ("Internal Mobility Rate", "12%", "15%", "-3%", "Down"),
        ("Diversity Index", "0.72", "0.75", "-0.03", "Down"),
        ("Engagement Score", "78%", "82%", "-4%", "Down"),
        ("Cost per Hire ($)", "4,850", "4,200", "+650", "Up"),
        ("New Hire Retention (90-day)", "91%", "95%", "-4%", "Down"),
    ]
    for metric, current, target, variance, trend in kpi_data:
        doc.add_paragraph(
            f'{metric}: Current Value = {current}, Target = {target}, '
            f'Variance = {variance}, Trend = {trend}'
        )

    # --- Section: Headcount by Department ---
    doc.add_heading('Headcount Analysis by Department', level=1)
    doc.add_paragraph('Department headcount breakdown as of December 31, 2025:')
    dept_data = [
        ("Engineering", 245, 12, 8),
        ("Sales", 180, 15, 10),
        ("Marketing", 95, 5, 3),
        ("Human Resources", 42, 3, 2),
        ("Finance", 68, 4, 3),
        ("Operations", 156, 8, 6),
        ("Legal", 28, 1, 1),
        ("Customer Support", 134, 10, 7),
        ("Product Management", 52, 4, 2),
        ("Research & Development", 88, 6, 4),
    ]
    for dept, hc, new_hires, departures in dept_data:
        doc.add_paragraph(
            f'{dept} - Headcount: {hc}, New Hires Q4: {new_hires}, Departures Q4: {departures}'
        )

    # --- Section: Headcount by Location ---
    doc.add_heading('Headcount Analysis by Location', level=1)
    doc.add_paragraph('Employee distribution across office locations:')
    loc_data = [
        ("San Francisco HQ", 420, "38.4%"),
        ("New York", 285, "26.1%"),
        ("Austin", 178, "16.3%"),
        ("Chicago", 124, "11.3%"),
        ("Remote", 81, "7.4%"),
    ]
    for loc, count, pct in loc_data:
        doc.add_paragraph(f'{loc} - Employees: {count}, Percentage: {pct}')

    # --- Section: Headcount by Job Level ---
    doc.add_heading('Headcount Analysis by Job Level', level=1)
    doc.add_paragraph('Distribution of employees across job levels:')
    level_data = [
        ("Executive (VP+)", 18, "$285,000"),
        ("Senior Director", 35, "$215,000"),
        ("Director", 72, "$175,000"),
        ("Senior Manager", 148, "$138,000"),
        ("Manager", 256, "$105,000"),
        ("Individual Contributor", 559, "$82,000"),
    ]
    for level, count, avg_comp in level_data:
        doc.add_paragraph(
            f'{level} - Count: {count}, Average Compensation: {avg_comp}'
        )

    # --- Section: Monthly Turnover ---
    doc.add_heading('Monthly Turnover Tracking', level=1)
    doc.add_paragraph(
        'Monthly voluntary and involuntary turnover rates for calendar year 2025:'
    )
    turnover_data = [
        ("January", "1.2%", "0.3%", "1.5%"),
        ("February", "1.0%", "0.2%", "1.2%"),
        ("March", "1.4%", "0.4%", "1.8%"),
        ("April", "1.1%", "0.3%", "1.4%"),
        ("May", "1.3%", "0.2%", "1.5%"),
        ("June", "1.5%", "0.5%", "2.0%"),
        ("July", "1.6%", "0.3%", "1.9%"),
        ("August", "1.2%", "0.4%", "1.6%"),
        ("September", "1.0%", "0.3%", "1.3%"),
        ("October", "1.3%", "0.4%", "1.7%"),
        ("November", "1.1%", "0.3%", "1.4%"),
        ("December", "0.9%", "0.2%", "1.1%"),
    ]
    for month, vol, invol, total in turnover_data:
        doc.add_paragraph(
            f'{month} 2025 - Voluntary: {vol}, Involuntary: {invol}, Total: {total}'
        )

    # --- Section: Cost per Hire ---
    doc.add_heading('Cost-per-Hire Analysis by Role Category', level=1)
    doc.add_paragraph(
        'Breakdown of average hiring costs by role category during 2025:'
    )
    cph_data = [
        ("Executive Leadership", "$28,500", 6, "$171,000"),
        ("Senior Technical", "$12,800", 24, "$307,200"),
        ("Mid-Level Technical", "$7,200", 45, "$324,000"),
        ("Junior Technical", "$4,500", 38, "$171,000"),
        ("Sales & Business Dev", "$6,100", 32, "$195,200"),
        ("Administrative & Support", "$3,200", 22, "$70,400"),
        ("Internship Program", "$1,800", 15, "$27,000"),
    ]
    for role, avg_cost, hires, total_cost in cph_data:
        doc.add_paragraph(
            f'{role} - Average Cost: {avg_cost}, Hires: {hires}, Total Cost: {total_cost}'
        )

    # --- Section: Time to Productivity ---
    doc.add_heading('Time-to-Productivity Metrics', level=1)
    doc.add_paragraph(
        'Average number of days for new hires to reach full productivity by department:'
    )
    ttp_data = [
        ("Engineering", 90, 85, "94%"),
        ("Sales", 120, 105, "88%"),
        ("Marketing", 75, 70, "93%"),
        ("Customer Support", 45, 42, "93%"),
        ("Finance", 60, 58, "97%"),
        ("Operations", 55, 50, "91%"),
        ("Legal", 80, 78, "98%"),
        ("Product Management", 85, 80, "94%"),
    ]
    for dept, target_days, actual_days, onboarding_completion in ttp_data:
        doc.add_paragraph(
            f'{dept} - Target Days: {target_days}, Actual Days: {actual_days}, '
            f'Onboarding Completion: {onboarding_completion}'
        )

    # --- Section: Engagement Score Trends ---
    doc.add_heading('Employee Engagement Score Trends', level=1)
    doc.add_paragraph(
        'Quarterly engagement survey results for 2024-2025:'
    )
    eng_data = [
        ("Q1 2024", 76, 82, 71, 88),
        ("Q2 2024", 78, 80, 73, 85),
        ("Q3 2024", 75, 79, 70, 86),
        ("Q4 2024", 77, 81, 72, 87),
        ("Q1 2025", 79, 83, 74, 89),
        ("Q2 2025", 80, 84, 75, 90),
        ("Q3 2025", 78, 82, 73, 88),
        ("Q4 2025", 78, 83, 72, 89),
    ]
    for period, overall, satisfaction, nps, recommend in eng_data:
        doc.add_paragraph(
            f'{period} - Overall Score: {overall}%, Job Satisfaction: {satisfaction}%, '
            f'eNPS: {nps}, Recommend to Friend: {recommend}%'
        )

    doc.add_paragraph('')
    doc.add_paragraph(
        'End of raw data. This information should be structured into formatted tables '
        'with proper captions, styling, and a master table of contents.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
