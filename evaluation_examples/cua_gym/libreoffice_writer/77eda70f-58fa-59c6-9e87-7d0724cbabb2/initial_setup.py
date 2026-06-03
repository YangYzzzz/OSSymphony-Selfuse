"""
Initial Setup: Configure document page numbering with Roman/Arabic sections
Task ID: writer_af_019
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_af_019'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# Namespace map
NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_footer_page_number(section, fmt='decimal'):
    """Add a simple page number in the footer of a section.
    fmt: 'decimal' for Arabic, 'lowerRoman' for lowercase Roman.
    """
    footer = section.footer
    footer.is_linked_to_previous = False
    # Clear existing footer content
    for p in footer.paragraphs:
        p.clear()

    fp = footer.paragraphs[0]
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Field: begin
    r1 = fp.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    # Field: instruction
    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)

    # Field: separate
    r3 = fp.add_run()
    fld_sep = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    r3._element.append(fld_sep)

    # Field: cached value
    r4 = fp.add_run('1')

    # Field: end
    r5 = fp.add_run()
    fld_end = r5._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r5._element.append(fld_end)


def set_page_number_format(section, fmt='decimal', start=None):
    """Set the page number format in the section properties.
    fmt: 'decimal', 'lowerRoman', 'upperRoman', etc.
    start: page number to start from (int) or None to continue.
    """
    sectPr = section._sectPr
    # Remove existing pgNumType if any
    for existing in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(existing)
    attrs = {qn('w:fmt'): fmt}
    if start is not None:
        attrs[qn('w:start')] = str(start)
    pgNumType = sectPr.makeelement(qn('w:pgNumType'), attrs)
    sectPr.append(pgNumType)


def create_initial():
    doc = Document()

    # --- Page setup for default section ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # =============================================
    # PREFACE SECTION (pages 1-4)
    # =============================================

    # --- Page 1: Title Page ---
    title = doc.add_heading('The Art of Modern Software Architecture', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(120)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_before = Pt(24)
    run = subtitle.add_run('A Comprehensive Guide to Building Scalable Systems')
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    author = doc.add_paragraph()
    author.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_before = Pt(48)
    run = author.add_run('Dr. Elena Vasquez')
    run.font.size = Pt(14)

    edition = doc.add_paragraph()
    edition.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    edition.paragraph_format.space_before = Pt(12)
    run = edition.add_run('Third Edition — 2025')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    publisher = doc.add_paragraph()
    publisher.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    publisher.paragraph_format.space_before = Pt(72)
    run = publisher.add_run('Nexus Technical Press\nSan Francisco, CA')
    run.font.size = Pt(11)

    # --- Page 2: Copyright / Legal ---
    doc.add_page_break()
    doc.add_paragraph('')  # spacer
    legal = doc.add_paragraph()
    legal.paragraph_format.space_before = Pt(200)
    legal.add_run('Copyright © 2025 Elena Vasquez. All rights reserved.').font.size = Pt(9)
    doc.add_paragraph('No part of this publication may be reproduced, distributed, or transmitted '
                      'in any form or by any means without the prior written permission of the publisher, '
                      'except in the case of brief quotations embodied in critical reviews.').runs[0].font.size = Pt(9)
    doc.add_paragraph('')
    isbn = doc.add_paragraph()
    isbn.add_run('ISBN 978-1-234567-89-0').font.size = Pt(9)
    doc.add_paragraph('Published by Nexus Technical Press').runs[0].font.size = Pt(9)
    doc.add_paragraph('First edition: 2019 | Second edition: 2022 | Third edition: 2025').runs[0].font.size = Pt(9)

    # --- Page 3: Table of Contents ---
    doc.add_page_break()
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.paragraph_format.space_after = Pt(18)

    toc_entries = [
        ('Part I: Foundations', ''),
        ('  Chapter 1: Principles of Software Architecture', '1'),
        ('  Chapter 2: Design Patterns Revisited', '15'),
        ('  Chapter 3: Architectural Decision Records', '35'),
        ('Part II: Distributed Systems', ''),
        ('  Chapter 4: Microservices Architecture', '52'),
        ('  Chapter 5: Event-Driven Design', '78'),
        ('  Chapter 6: Service Mesh and API Gateways', '101'),
        ('Part III: Data Architecture', ''),
        ('  Chapter 7: Database Selection Strategies', '125'),
        ('  Chapter 8: CQRS and Event Sourcing', '148'),
        ('  Chapter 9: Data Pipeline Architecture', '172'),
        ('Part IV: Operations', ''),
        ('  Chapter 10: Infrastructure as Code', '195'),
        ('  Chapter 11: Observability and Monitoring', '218'),
        ('  Chapter 12: Chaos Engineering', '240'),
        ('Appendix A: Architecture Review Checklist', '258'),
        ('Appendix B: Technology Radar', '265'),
        ('Index', '275'),
    ]
    for entry, page in toc_entries:
        p = doc.add_paragraph()
        if not entry.startswith('  '):
            run = p.add_run(entry)
            run.bold = True
            run.font.size = Pt(12)
        else:
            run = p.add_run(entry)
            run.font.size = Pt(11)
        if page:
            tab_run = p.add_run(f'\t{page}')
            tab_run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    # --- Page 4: Preface ---
    doc.add_page_break()
    doc.add_heading('Preface', level=1)
    preface_text = [
        'When I wrote the first edition of this book in 2019, the landscape of software '
        'architecture was already undergoing dramatic transformation. Microservices were moving '
        'from buzzword to mainstream practice. Kubernetes was becoming the de facto standard for '
        'container orchestration. Cloud-native architectures were no longer optional.',

        'Six years later, the pace of change has only accelerated. The rise of AI-driven development, '
        'edge computing, and serverless architectures has created new paradigms that demand fresh '
        'thinking about how we design and build software systems.',

        'This third edition reflects these changes while staying true to the core principles that '
        'have guided architects for decades. Good architecture is not about following the latest '
        'trends — it is about making deliberate decisions that balance competing concerns.',

        'I am deeply grateful to my colleagues at Nexus Labs, particularly Dr. James Chen and '
        'Prof. Amara Okafor, whose insights on distributed systems and data architecture have '
        'greatly enriched this edition. Special thanks to the technical reviewers: Maria Santos, '
        'Raj Patel, and Tomas Eriksson.',

        '— Elena Vasquez, San Francisco, January 2025',
    ]
    for text in preface_text:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)

    # =============================================
    # SECTION BREAK: Preface -> Main Content
    # Use python-docx add_section for proper section management
    # =============================================
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    # =============================================
    # MAIN CONTENT (pages 5-60)
    # =============================================

    # Chapter content to fill ~56 pages
    chapters = [
        {
            'title': 'Chapter 1: Principles of Software Architecture',
            'sections': [
                ('What is Software Architecture?',
                 'Software architecture refers to the fundamental structures of a software system and the '
                 'discipline of creating such structures and systems. Each structure comprises software elements, '
                 'relations among them, and properties of both elements and relations. The architecture of a '
                 'software system is a metaphor, analogous to the architecture of a building.'),
                ('Architectural Styles and Patterns',
                 'An architectural style defines a family of systems in terms of a pattern of structural '
                 'organization. More specifically, an architectural style determines the vocabulary of '
                 'components and connectors that can be used in instances of that style, together with a set '
                 'of constraints on how they can be combined. Layered architecture, for instance, organizes '
                 'the system into layers with each layer providing services to the layer above it.'),
                ('Quality Attributes',
                 'Quality attributes are the overall factors that affect runtime behavior, system design, and '
                 'user experience. They represent areas of concern that have the potential for application-wide '
                 'impact across layers and tiers. Some of these attributes are related to the overall system '
                 'design, while others are specific to runtime, design time, or user-centric issues. '
                 'Performance, scalability, reliability, availability, security, and maintainability are the '
                 'most commonly cited quality attributes in modern systems.'),
                ('Trade-off Analysis',
                 'Every architectural decision involves trade-offs. Improving performance may reduce '
                 'maintainability. Increasing security may impact usability. The Architecture Tradeoff '
                 'Analysis Method (ATAM) provides a structured approach to evaluating these trade-offs '
                 'by examining architectural decisions in the context of quality attribute requirements.'),
            ]
        },
        {
            'title': 'Chapter 2: Design Patterns Revisited',
            'sections': [
                ('The Evolution of Design Patterns',
                 'Since the Gang of Four published their seminal work in 1994, design patterns have become '
                 'a cornerstone of software engineering education and practice. However, the landscape has '
                 'shifted considerably. Many patterns that were essential in C++ and Java have become less '
                 'relevant with modern language features, while new patterns have emerged to address '
                 'distributed systems and cloud-native architectures.'),
                ('Creational Patterns in Modern Context',
                 'Factory patterns remain relevant but have evolved. Dependency injection containers have '
                 'largely replaced manual factory implementations. The Builder pattern has found new life in '
                 'fluent APIs and configuration objects. Abstract Factory is now often implemented through '
                 'service provider interfaces and plugin architectures.'),
                ('Structural Patterns for Microservices',
                 'The Adapter pattern is essential for integrating legacy systems with modern APIs. The '
                 'Facade pattern manifests as API gateways and Backend-for-Frontend (BFF) services. The '
                 'Proxy pattern underpins service meshes and sidecar containers. The Decorator pattern '
                 'enables middleware chains in HTTP frameworks and message processing pipelines.'),
                ('Behavioral Patterns in Event-Driven Systems',
                 'The Observer pattern has evolved into publish-subscribe systems and event streaming '
                 'platforms. The Strategy pattern enables runtime algorithm selection in feature flag '
                 'systems. The Chain of Responsibility pattern structures request processing in middleware '
                 'stacks and approval workflows. The Command pattern enables event sourcing and CQRS.'),
            ]
        },
        {
            'title': 'Chapter 3: Architectural Decision Records',
            'sections': [
                ('Why Document Decisions?',
                 'Architecture decisions are among the most important technical decisions in a project. '
                 'They are expensive to change and have long-lasting effects. Yet in many organizations, '
                 'these decisions exist only in the memories of senior engineers or buried in meeting notes. '
                 'Architectural Decision Records (ADRs) provide a lightweight method to capture these '
                 'decisions and their rationale.'),
                ('ADR Format and Best Practices',
                 'An ADR captures a single decision in a concise format: Title, Status (Proposed, Accepted, '
                 'Deprecated, Superseded), Context (the forces at play), Decision (the response to these '
                 'forces), and Consequences (the resulting context after the decision). Each ADR should be '
                 'numbered sequentially and stored in version control alongside the code it governs.'),
                ('Decision Governance',
                 'Effective decision governance requires clear ownership, review processes, and periodic '
                 'reassessment. Architecture Review Boards (ARBs) should review significant decisions, but '
                 'the process should not become a bottleneck. Lightweight governance, where teams document '
                 'decisions and flag those exceeding defined thresholds for review, balances autonomy with '
                 'organizational coherence.'),
            ]
        },
        {
            'title': 'Chapter 4: Microservices Architecture',
            'sections': [
                ('Decomposition Strategies',
                 'The transition from monolith to microservices requires careful decomposition. Domain-Driven '
                 'Design provides the most effective framework through bounded contexts. Each bounded context '
                 'becomes a candidate microservice boundary. The Strangler Fig pattern enables incremental '
                 'migration by gradually routing functionality from the monolith to new services.'),
                ('Inter-Service Communication',
                 'Synchronous communication via REST or gRPC provides simplicity and immediate feedback. '
                 'Asynchronous communication via message queues or event streams provides better decoupling '
                 'and resilience. Most real-world systems use a mix: synchronous for queries and commands '
                 'requiring immediate confirmation, asynchronous for notifications and background processing.'),
                ('Data Management in Microservices',
                 'The database-per-service pattern ensures loose coupling but introduces the challenge of '
                 'maintaining data consistency across services. The Saga pattern coordinates distributed '
                 'transactions through a sequence of local transactions. The CQRS pattern separates read '
                 'and write models, enabling optimized data access patterns for different use cases.'),
                ('Service Discovery and Load Balancing',
                 'In a dynamic microservices environment, services must find each other without hard-coded '
                 'addresses. Client-side discovery uses a service registry that clients query directly. '
                 'Server-side discovery delegates routing to a load balancer or API gateway. Platform-based '
                 'discovery leverages orchestration platforms like Kubernetes with built-in DNS and service '
                 'abstractions.'),
            ]
        },
        {
            'title': 'Chapter 5: Event-Driven Design',
            'sections': [
                ('Event Sourcing Fundamentals',
                 'Event sourcing stores the state of a business entity as a sequence of state-changing '
                 'events. Instead of persisting the current state, the application stores every change as '
                 'an event in an append-only log. The current state is reconstructed by replaying events. '
                 'This provides a complete audit trail and enables temporal queries.'),
                ('Event Streaming Platforms',
                 'Apache Kafka has become the dominant event streaming platform, providing durable, '
                 'high-throughput, distributed commit logs. Kafka Streams enables real-time stream processing '
                 'without external frameworks. Apache Pulsar offers multi-tenancy and geo-replication as '
                 'first-class features. Amazon Kinesis provides a managed alternative for AWS-centric '
                 'architectures.'),
                ('Designing Event Schemas',
                 'Event schema design is critical for long-term system evolution. Events should be '
                 'self-describing and include sufficient context to be processed independently. Schema '
                 'registries enforce compatibility rules (backward, forward, full) to prevent breaking '
                 'changes. Avro and Protocol Buffers provide efficient serialization with schema evolution '
                 'support.'),
            ]
        },
        {
            'title': 'Chapter 6: Service Mesh and API Gateways',
            'sections': [
                ('The Service Mesh Architecture',
                 'A service mesh provides a dedicated infrastructure layer for handling service-to-service '
                 'communication. It manages traffic routing, load balancing, authentication, authorization, '
                 'and observability through sidecar proxies deployed alongside each service instance. '
                 'Istio and Linkerd are the leading open-source service mesh implementations.'),
                ('API Gateway Patterns',
                 'API gateways serve as the entry point for external clients, providing request routing, '
                 'protocol translation, rate limiting, and authentication. The Backend-for-Frontend pattern '
                 'creates specialized gateways for each client type (web, mobile, IoT), optimizing the API '
                 'surface for each use case. Kong, Ambassador, and AWS API Gateway are popular choices.'),
            ]
        },
        {
            'title': 'Chapter 7: Database Selection Strategies',
            'sections': [
                ('The Polyglot Persistence Approach',
                 'Modern applications rarely rely on a single database technology. Polyglot persistence '
                 'uses different storage technologies for different data access patterns. Relational databases '
                 'for transactional data, document stores for flexible schemas, graph databases for '
                 'relationship-heavy queries, and time-series databases for metrics and IoT data.'),
                ('CAP Theorem and Its Practical Implications',
                 'The CAP theorem states that a distributed data store can provide only two of three '
                 'guarantees: Consistency, Availability, and Partition tolerance. In practice, partition '
                 'tolerance is non-negotiable in distributed systems, so the real choice is between '
                 'consistency and availability during network partitions. Different parts of a system may '
                 'make different choices based on their requirements.'),
                ('NewSQL and Distributed SQL',
                 'NewSQL databases like CockroachDB, TiDB, and YugabyteDB combine the ACID guarantees '
                 'of traditional relational databases with the horizontal scalability of NoSQL systems. '
                 'They use distributed consensus protocols (Raft, Paxos) to maintain consistency across '
                 'nodes while supporting standard SQL interfaces.'),
            ]
        },
        {
            'title': 'Chapter 8: CQRS and Event Sourcing',
            'sections': [
                ('Command Query Responsibility Segregation',
                 'CQRS separates read and write operations into distinct models. The write model handles '
                 'commands that change state, optimized for validation and business rule enforcement. The '
                 'read model handles queries, optimized for specific view requirements. This separation '
                 'enables independent scaling and optimization of each concern.'),
                ('Implementing Event Sourcing with CQRS',
                 'When combined with event sourcing, CQRS becomes particularly powerful. Commands produce '
                 'events stored in an event store. Read models (projections) are built by processing event '
                 'streams, creating denormalized views optimized for specific query patterns. Multiple '
                 'projections can exist for the same event stream, each serving different needs.'),
            ]
        },
        {
            'title': 'Chapter 9: Data Pipeline Architecture',
            'sections': [
                ('Batch vs. Stream Processing',
                 'Traditional batch processing (MapReduce, Spark) processes bounded datasets in scheduled '
                 'intervals. Stream processing (Flink, Kafka Streams) processes unbounded data in real-time. '
                 'The Lambda architecture combines both, maintaining a batch layer for historical accuracy '
                 'and a speed layer for real-time updates. The Kappa architecture simplifies this by using '
                 'stream processing for both real-time and historical data.'),
                ('Data Quality and Observability',
                 'Data quality in pipelines requires validation at ingestion, transformation, and output '
                 'stages. Great Expectations provides a framework for defining and testing data quality '
                 'rules. Data observability platforms like Monte Carlo and Bigeye detect anomalies in data '
                 'freshness, volume, schema changes, and distribution shifts.'),
            ]
        },
        {
            'title': 'Chapter 10: Infrastructure as Code',
            'sections': [
                ('Declarative vs. Imperative Approaches',
                 'Infrastructure as Code (IaC) enables reproducible, version-controlled infrastructure '
                 'provisioning. Declarative tools (Terraform, CloudFormation) define desired end-state, and '
                 'the tool determines the necessary changes. Imperative tools (Pulumi, AWS CDK) use '
                 'programming languages to define infrastructure, enabling loops, conditionals, and '
                 'abstractions. Both approaches have trade-offs in complexity, debugging, and team adoption.'),
                ('GitOps and Continuous Infrastructure',
                 'GitOps extends IaC by using Git as the single source of truth for both application and '
                 'infrastructure state. Pull-based reconciliation (ArgoCD, Flux) automatically detects '
                 'drift between the desired state in Git and the actual state in the cluster, applying '
                 'corrections without manual intervention. This provides an audit trail, enables rollback, '
                 'and enforces review processes for all infrastructure changes.'),
                ('Platform Engineering',
                 'Platform engineering builds internal developer platforms that abstract infrastructure '
                 'complexity. Developer portals (Backstage), self-service provisioning, and golden paths '
                 'reduce cognitive load on application teams. The platform team provides curated, '
                 'opinionated tooling that balances developer autonomy with organizational standards '
                 'for security, compliance, and operational excellence.'),
            ]
        },
        {
            'title': 'Chapter 11: Observability and Monitoring',
            'sections': [
                ('The Three Pillars of Observability',
                 'Observability in distributed systems rests on three pillars: logs (discrete events), '
                 'metrics (aggregated measurements), and traces (request flows across services). OpenTelemetry '
                 'provides a vendor-neutral standard for collecting all three signal types. Effective '
                 'observability enables teams to understand system behavior without deploying new code.'),
                ('SLOs, SLIs, and Error Budgets',
                 'Service Level Objectives (SLOs) define reliability targets based on Service Level '
                 'Indicators (SLIs) — measurable aspects of service behavior like latency and error rate. '
                 'Error budgets quantify acceptable unreliability: if the SLO is 99.9% availability, the '
                 'error budget is 0.1% downtime per period. Teams spend error budget on feature velocity '
                 'and reclaim it by investing in reliability when the budget runs low.'),
            ]
        },
        {
            'title': 'Chapter 12: Chaos Engineering',
            'sections': [
                ('Principles of Chaos Engineering',
                 'Chaos engineering is the discipline of experimenting on a system to build confidence in '
                 'its capability to withstand turbulent conditions in production. It involves forming '
                 'hypotheses about steady-state behavior, introducing variables that reflect real-world '
                 'events (server failures, network partitions, resource exhaustion), and observing the '
                 'difference between the hypothesis and actual behavior.'),
                ('Implementing Chaos Experiments',
                 'Start with game days — planned exercises where teams manually introduce failures. '
                 'Progress to automated chaos tools like Chaos Monkey, Litmus, and Gremlin. Always '
                 'have a rollback plan. Start in staging environments before moving to production. '
                 'Measure blast radius and implement circuit breakers to limit unexpected damage.'),
                ('Building a Chaos Culture',
                 'Successful chaos engineering requires organizational support. Blameless postmortems '
                 'create psychological safety for experimentation. Start small — a single service, a '
                 'single failure mode — and build confidence gradually. Share results widely to '
                 'demonstrate value and build institutional knowledge about system resilience.'),
            ]
        },
    ]

    for ch_idx, chapter in enumerate(chapters):
        doc.add_heading(chapter['title'], level=1)
        for sec_title, sec_text in chapter['sections']:
            doc.add_heading(sec_title, level=2)
            # Add multiple paragraphs to fill pages
            p = doc.add_paragraph(sec_text)
            p.paragraph_format.space_after = Pt(6)
            # Add extra filler to create more pages
            elaboration = (
                f'This topic has been extensively studied in the academic literature and '
                f'validated through practical application in enterprise environments. '
                f'Organizations adopting these practices report significant improvements in '
                f'system reliability, team productivity, and time-to-market for new features. '
                f'The key insight is that architecture decisions should be driven by measurable '
                f'quality attributes rather than technology trends. Empirical evidence from '
                f'industry surveys consistently shows that teams investing in architectural '
                f'practices early in the development lifecycle achieve better outcomes than '
                f'those deferring these decisions.'
            )
            p2 = doc.add_paragraph(elaboration)
            p2.paragraph_format.space_after = Pt(6)

            case_study = (
                f'Consider the case of a major financial services company that migrated its '
                f'core trading platform from a monolithic architecture to a distributed system. '
                f'The migration took 18 months and involved restructuring 2.3 million lines of '
                f'code across 47 services. Performance improved by 340%, deployment frequency '
                f'increased from monthly to daily, and incident resolution time decreased from '
                f'4 hours to 15 minutes on average. These results demonstrate the tangible '
                f'benefits of systematic architectural improvement.'
            )
            p3 = doc.add_paragraph(case_study)
            p3.paragraph_format.space_after = Pt(12)

    # Add appendices to fill remaining pages
    doc.add_heading('Appendix A: Architecture Review Checklist', level=1)
    checklist_items = [
        'System context and boundary definition',
        'Quality attribute requirements and priorities',
        'Component decomposition and responsibility assignment',
        'Inter-component communication patterns',
        'Data management strategy and consistency model',
        'Security architecture and threat model',
        'Deployment topology and infrastructure requirements',
        'Monitoring, alerting, and incident response plan',
        'Disaster recovery and business continuity',
        'Performance benchmarks and capacity planning',
        'API design standards and versioning strategy',
        'Technical debt assessment and remediation plan',
    ]
    for item in checklist_items:
        p = doc.add_paragraph(f'☐ {item}', style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        detail = doc.add_paragraph(
            f'Verify that the {item.lower()} has been documented, reviewed by stakeholders, '
            f'and validated against organizational standards. Include relevant metrics, '
            f'acceptance criteria, and known limitations or risks.'
        )
        detail.paragraph_format.space_after = Pt(8)

    doc.add_heading('Appendix B: Technology Radar', level=1)
    radar_text = (
        'This technology radar captures the current assessment of key technologies relevant '
        'to software architecture. Technologies are classified into four rings: Adopt (proven '
        'and recommended), Trial (worth pursuing in low-risk projects), Assess (worth exploring '
        'to understand applicability), and Hold (proceed with caution).'
    )
    doc.add_paragraph(radar_text)

    adopt_items = [
        'Kubernetes', 'Terraform', 'GraphQL', 'OpenTelemetry', 'ArgoCD',
        'PostgreSQL', 'Redis', 'Apache Kafka', 'gRPC', 'Backstage',
    ]
    trial_items = [
        'WebAssembly (server-side)', 'CockroachDB', 'Dapr', 'Temporal.io',
        'eBPF-based networking', 'NATS JetStream',
    ]
    assess_items = [
        'AI-assisted code generation', 'Confidential computing', 'RISC-V servers',
        'Decentralized identity', 'Quantum-resistant cryptography',
    ]
    hold_items = [
        'Self-hosted CI/CD (prefer managed)', 'XML-based configuration',
        'Synchronous inter-service calls for all communication',
        'Monolithic deployments for greenfield projects',
    ]

    for ring_name, items in [('Adopt', adopt_items), ('Trial', trial_items),
                              ('Assess', assess_items), ('Hold', hold_items)]:
        doc.add_heading(ring_name, level=2)
        for item in items:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            strength = 'strong' if ring_name in ('Adopt', 'Trial') else 'emerging'
            advice = ''
            if ring_name == 'Adopt':
                advice = 'Recommended for broad adoption.'
            elif ring_name == 'Trial':
                advice = 'Consider for pilot projects.'
            elif ring_name == 'Assess':
                advice = 'Monitor developments closely.'
            elif ring_name == 'Hold':
                advice = 'Evaluate alternatives before proceeding.'
            desc = doc.add_paragraph(
                '{} has demonstrated {} value in production environments. {}'.format(item, strength, advice)
            )
            desc.paragraph_format.space_after = Pt(4)

    # Add index placeholder to fill more pages
    doc.add_heading('Index', level=1)
    index_entries = [
        'API Gateway, 101-108',
        'Architectural Decision Records (ADR), 35-51',
        'Availability, 130-133',
        'Batch processing, 172-178',
        'Bounded context, 52-58',
        'CAP theorem, 133-138',
        'CQRS, 148-162',
        'Chaos engineering, 240-257',
        'Circuit breaker, 82-86',
        'Consistency models, 134-140',
        'Container orchestration, 195-200',
        'Data pipeline, 172-194',
        'Decomposition strategies, 52-60',
        'Dependency injection, 18-22',
        'Design patterns, 15-34',
        'Domain-driven design, 52-58',
        'Error budget, 230-235',
        'Event sourcing, 78-92, 148-162',
        'Event streaming, 85-92',
        'Fault tolerance, 240-252',
        'GitOps, 205-212',
        'Graph database, 128-130',
        'gRPC, 65-70',
        'Infrastructure as Code, 195-217',
        'Kafka, 85-92',
        'Kubernetes, 195-204',
        'Lambda architecture, 175-178',
        'Load balancing, 68-72',
        'Microservices, 52-77',
        'Monitoring, 218-239',
        'NewSQL, 140-145',
        'Observability, 218-235',
        'OpenTelemetry, 222-228',
        'Platform engineering, 212-217',
        'Polyglot persistence, 125-128',
        'Quality attributes, 8-14',
        'REST, 62-65',
        'Saga pattern, 60-62',
        'Schema evolution, 92-97',
        'Service discovery, 68-72',
        'Service Level Objectives (SLO), 230-235',
        'Service mesh, 101-112',
        'Stream processing, 178-185',
        'Terraform, 197-204',
        'Trade-off analysis, 12-14',
    ]
    for entry in index_entries:
        p = doc.add_paragraph(entry)
        p.paragraph_format.space_after = Pt(1)
        p.runs[0].font.size = Pt(9)

    # Now configure the sections:
    # Section 0 = preface (pages 1-4): Arabic numbering starting from 1
    # Section 1 = main content (pages 5+): Arabic numbering continuing
    # ALL pages use standard Arabic (decimal) numbering — this is the INITIAL state
    # The task is to CHANGE this to Roman for preface and restart Arabic for main

    # Section 0 (preface): Arabic from 1
    sect0 = doc.sections[0]
    set_page_number_format(sect0, fmt='decimal', start=1)
    add_footer_page_number(sect0, fmt='decimal')

    # Section 1 (main content): Arabic continuing from 5
    sect1 = doc.sections[1]
    set_page_number_format(sect1, fmt='decimal', start=5)
    add_footer_page_number(sect1, fmt='decimal')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
