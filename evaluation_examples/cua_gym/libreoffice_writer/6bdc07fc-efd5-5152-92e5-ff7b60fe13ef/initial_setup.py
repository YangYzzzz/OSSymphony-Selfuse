"""
Initial Setup: User Manual with tracked changes
Task ID: writer_rm_019
Domain: libreoffice_writer

Creates an 8-page User Manual document with 7 tracked changes.
The first tracked change is in paragraph 4: 'click' -> 'select'.
"""

import os
import shlex
import subprocess
import time
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_019'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

AUTHOR = 'Emily Watson'
DATE = '2026-03-28T14:30:00Z'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_run_element(text, bold=False, italic=False, font_name='Calibri', font_size=11):
    """Create a w:r element with optional formatting."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if font_name:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)
    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size * 2))  # half-points
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(font_size * 2))
        rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if italic:
        i = OxmlElement('w:i')
        rPr.append(i)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def add_tracked_change_replace(para_element, rev_id, old_text, new_text, author=AUTHOR, date=DATE):
    """
    Replace 'old_text' with a tracked change (del old + ins new) in the paragraph.
    Finds the run containing old_text, splits it, and inserts del/ins elements.
    """
    ns = {'w': W_NS}
    runs = para_element.findall('.//w:r', ns)

    for run in runs:
        t_elem = run.find('w:t', ns)
        if t_elem is None or t_elem.text is None:
            continue
        if old_text not in t_elem.text:
            continue

        # Found the run containing old_text
        full_text = t_elem.text
        idx = full_text.index(old_text)
        before = full_text[:idx]
        after = full_text[idx + len(old_text):]

        # Get run properties for cloning
        rPr = run.find('w:rPr', ns)
        rPr_xml = etree.tostring(rPr) if rPr is not None else None

        parent = run.getparent()
        run_index = list(parent).index(run)

        # Remove original run
        parent.remove(run)

        insert_pos = run_index

        # Before text run
        if before:
            r_before = make_run_element(before)
            if rPr_xml is not None:
                r_before.remove(r_before.find(qn('w:rPr')))
                r_before.insert(0, etree.fromstring(rPr_xml))
            parent.insert(insert_pos, r_before)
            insert_pos += 1

        # Deletion element
        w_del = OxmlElement('w:del')
        w_del.set(qn('w:id'), str(rev_id))
        w_del.set(qn('w:author'), author)
        w_del.set(qn('w:date'), date)

        del_run = OxmlElement('w:r')
        if rPr_xml is not None:
            del_run.insert(0, etree.fromstring(rPr_xml))
        del_text = OxmlElement('w:delText')
        del_text.set(qn('xml:space'), 'preserve')
        del_text.text = old_text
        del_run.append(del_text)
        w_del.append(del_run)
        parent.insert(insert_pos, w_del)
        insert_pos += 1

        # Insertion element
        w_ins = OxmlElement('w:ins')
        w_ins.set(qn('w:id'), str(rev_id + 1))
        w_ins.set(qn('w:author'), author)
        w_ins.set(qn('w:date'), date)

        ins_run = make_run_element(new_text)
        if rPr_xml is not None:
            ins_run.remove(ins_run.find(qn('w:rPr')))
            ins_run.insert(0, etree.fromstring(rPr_xml))
        w_ins.append(ins_run)
        parent.insert(insert_pos, w_ins)
        insert_pos += 1

        # After text run
        if after:
            r_after = make_run_element(after)
            if rPr_xml is not None:
                r_after.remove(r_after.find(qn('w:rPr')))
                r_after.insert(0, etree.fromstring(rPr_xml))
            parent.insert(insert_pos, r_after)

        return True
    return False


def add_tracked_insertion(para_element, rev_id, after_text, inserted_text, author=AUTHOR, date=DATE):
    """Insert new tracked text after a given phrase in the paragraph."""
    ns = {'w': W_NS}
    runs = para_element.findall('.//w:r', ns)

    for run in runs:
        t_elem = run.find('w:t', ns)
        if t_elem is None or t_elem.text is None:
            continue
        if after_text not in t_elem.text:
            continue

        full_text = t_elem.text
        idx = full_text.index(after_text) + len(after_text)
        before = full_text[:idx]
        after = full_text[idx:]

        rPr = run.find('w:rPr', ns)
        rPr_xml = etree.tostring(rPr) if rPr is not None else None

        parent = run.getparent()
        run_index = list(parent).index(run)
        parent.remove(run)
        insert_pos = run_index

        # Before text
        r_before = make_run_element(before)
        if rPr_xml is not None:
            r_before.remove(r_before.find(qn('w:rPr')))
            r_before.insert(0, etree.fromstring(rPr_xml))
        parent.insert(insert_pos, r_before)
        insert_pos += 1

        # Insertion
        w_ins = OxmlElement('w:ins')
        w_ins.set(qn('w:id'), str(rev_id))
        w_ins.set(qn('w:author'), author)
        w_ins.set(qn('w:date'), date)
        ins_run = make_run_element(inserted_text)
        if rPr_xml is not None:
            ins_run.remove(ins_run.find(qn('w:rPr')))
            ins_run.insert(0, etree.fromstring(rPr_xml))
        w_ins.append(ins_run)
        parent.insert(insert_pos, w_ins)
        insert_pos += 1

        # After text
        if after:
            r_after = make_run_element(after)
            if rPr_xml is not None:
                r_after.remove(r_after.find(qn('w:rPr')))
                r_after.insert(0, etree.fromstring(rPr_xml))
            parent.insert(insert_pos, r_after)

        return True
    return False


def add_tracked_deletion(para_element, rev_id, deleted_text, author=AUTHOR, date=DATE):
    """Mark existing text as deleted (tracked deletion)."""
    ns = {'w': W_NS}
    runs = para_element.findall('.//w:r', ns)

    for run in runs:
        t_elem = run.find('w:t', ns)
        if t_elem is None or t_elem.text is None:
            continue
        if deleted_text not in t_elem.text:
            continue

        full_text = t_elem.text
        idx = full_text.index(deleted_text)
        before = full_text[:idx]
        after = full_text[idx + len(deleted_text):]

        rPr = run.find('w:rPr', ns)
        rPr_xml = etree.tostring(rPr) if rPr is not None else None

        parent = run.getparent()
        run_index = list(parent).index(run)
        parent.remove(run)
        insert_pos = run_index

        if before:
            r_before = make_run_element(before)
            if rPr_xml is not None:
                r_before.remove(r_before.find(qn('w:rPr')))
                r_before.insert(0, etree.fromstring(rPr_xml))
            parent.insert(insert_pos, r_before)
            insert_pos += 1

        w_del = OxmlElement('w:del')
        w_del.set(qn('w:id'), str(rev_id))
        w_del.set(qn('w:author'), author)
        w_del.set(qn('w:date'), date)
        del_run = OxmlElement('w:r')
        if rPr_xml is not None:
            del_run.insert(0, etree.fromstring(rPr_xml))
        del_text = OxmlElement('w:delText')
        del_text.set(qn('xml:space'), 'preserve')
        del_text.text = deleted_text
        del_run.append(del_text)
        w_del.append(del_run)
        parent.insert(insert_pos, w_del)
        insert_pos += 1

        if after:
            r_after = make_run_element(after)
            if rPr_xml is not None:
                r_after.remove(r_after.find(qn('w:rPr')))
                r_after.insert(0, etree.fromstring(rPr_xml))
            parent.insert(insert_pos, r_after)

        return True
    return False


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ==================== PAGE 1 ====================
    # Paragraph 1: Title
    title = doc.add_heading('TaskFlow Pro - User Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Paragraph 2: Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Version 3.2 — March 2026')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Paragraph 3: Intro heading
    doc.add_heading('1. Introduction', level=1)

    # Paragraph 4: First paragraph of intro — THIS IS WHERE TRACKED CHANGE #1 GOES
    # "click" -> "select" tracked change
    p4 = doc.add_paragraph()
    p4.add_run(
        'TaskFlow Pro is a comprehensive project management solution designed for teams '
        'of all sizes. To get started with your first project, click the "New Project" '
        'button on the dashboard. This will open the project creation wizard, which guides '
        'you through setting up your workspace, inviting team members, and configuring '
        'notification preferences.'
    )

    # Paragraph 5
    p5 = doc.add_paragraph()
    p5.add_run(
        'The application supports multiple workflows including Kanban boards, Gantt charts, '
        'and sprint planning. Each workflow can be customized to match your team\'s specific '
        'needs. You can switch between views at any time using the toolbar at the top of '
        'the screen.'
    )

    # Paragraph 6
    p6 = doc.add_paragraph()
    p6.add_run(
        'TaskFlow Pro integrates seamlessly with popular tools such as Slack, Microsoft Teams, '
        'GitHub, and Jira. These integrations allow you to receive notifications, create tasks '
        'from messages, and synchronize project data across platforms without manual intervention.'
    )

    # ==================== PAGE 2 ====================
    doc.add_heading('2. Dashboard Overview', level=1)

    p_dash1 = doc.add_paragraph()
    p_dash1.add_run(
        'The main dashboard provides a consolidated view of all your active projects, recent '
        'activity, and upcoming deadlines. The left sidebar contains navigation links to your '
        'projects, team directory, and settings. The central panel displays project cards that '
        'can be sorted by priority, due date, or last modified.'
    )

    p_dash2 = doc.add_paragraph()
    p_dash2.add_run(
        'Each project card shows a progress indicator, the number of open tasks, and the '
        'assigned team lead. Hovering over a card reveals a quick-action menu where you can '
        'archive, duplicate, or share the project. The dashboard also features a global search '
        'bar that indexes all tasks, comments, and attachments across every project.'
    )

    p_dash3 = doc.add_paragraph()
    p_dash3.add_run(
        'Widgets on the dashboard can be rearranged by dragging them to your preferred position. '
        'Available widgets include a calendar view, a burndown chart, a team workload heatmap, '
        'and a recent files panel. Custom widgets can be created using the built-in widget editor, '
        'which supports HTML and JavaScript for advanced layouts.'
    )

    doc.add_page_break()

    # ==================== PAGE 3 ====================
    doc.add_heading('3. Task Management', level=1)

    p_task1 = doc.add_paragraph()
    p_task1.add_run(
        'Tasks are the core building blocks of any project in TaskFlow Pro. Each task has a title, '
        'description, assignee, priority level, and due date. You can create subtasks to break '
        'larger items into manageable pieces. Task dependencies allow you to specify that one task '
        'must be completed before another can begin.'
    )

    p_task2 = doc.add_paragraph()
    p_task2.add_run(
        'Labels and tags provide flexible categorization for your tasks. You can create custom '
        'label sets per project or use global labels that span your entire organization. The '
        'filter panel on the left side of the task board lets you narrow results by assignee, '
        'label, priority, or date range. Saved filters can be bookmarked for quick access.'
    )

    p_task3 = doc.add_paragraph()
    p_task3.add_run(
        'Every task maintains a complete history of changes including status updates, comment '
        'additions, and file attachments. The activity feed on each task card shows who made '
        'changes and when. Task templates let you pre-define common task structures so you '
        'can create new items quickly without entering repetitive details.'
    )

    p_task4 = doc.add_paragraph()
    p_task4.add_run(
        'Recurring tasks can be configured to automatically regenerate on a daily, weekly, or '
        'monthly schedule. When a recurring task is completed, a new instance is created with '
        'the same properties and assigned to the original owner. This feature is particularly '
        'useful for maintenance checklists and regular reporting tasks.'
    )

    doc.add_page_break()

    # ==================== PAGE 4 ====================
    doc.add_heading('4. Team Collaboration', level=1)

    p_collab1 = doc.add_paragraph()
    p_collab1.add_run(
        'Real-time collaboration is at the heart of TaskFlow Pro. Multiple team members can '
        'edit the same task simultaneously, with changes appearing instantly for all participants. '
        'The built-in chat feature allows you to discuss tasks directly within the context '
        'of the project, eliminating the need for external messaging tools.'
    )

    p_collab2 = doc.add_paragraph()
    p_collab2.add_run(
        'Mentions using the @ symbol notify specific team members and draw their attention '
        'to important updates. You can mention individuals or entire teams. Notification '
        'preferences can be configured per project to control the volume and type of alerts '
        'each member receives, whether by email, push notification, or in-app badge.'
    )

    p_collab3 = doc.add_paragraph()
    p_collab3.add_run(
        'File sharing within TaskFlow Pro supports documents up to 250 MB per upload. '
        'Version control tracks every revision of shared files, allowing team members to '
        'revert to previous versions if needed. The preview panel renders common file '
        'formats including PDFs, images, spreadsheets, and presentation slides.'
    )

    doc.add_page_break()

    # ==================== PAGE 5 ====================
    doc.add_heading('5. Reporting and Analytics', level=1)

    p_report1 = doc.add_paragraph()
    p_report1.add_run(
        'TaskFlow Pro includes a powerful reporting engine that generates insights from your '
        'project data. Standard reports include velocity charts, burndown analysis, team '
        'utilization rates, and cycle time distributions. Reports can be exported to PDF, '
        'Excel, or CSV format for sharing with stakeholders who may not have application access.'
    )

    p_report2 = doc.add_paragraph()
    p_report2.add_run(
        'Custom reports can be built using the visual report designer. You can drag and drop '
        'data fields, apply filters, and choose from a variety of chart types including bar, '
        'line, pie, and scatter plots. Scheduled reports can be emailed to specified recipients '
        'on a recurring basis, ensuring that leadership always has up-to-date project metrics.'
    )

    p_report3 = doc.add_paragraph()
    p_report3.add_run(
        'The analytics dashboard provides real-time KPIs for each project. Metrics such as '
        'task completion rate, average resolution time, and workload distribution are displayed '
        'in interactive charts. You can drill down into any metric to see the underlying data '
        'and identify trends or bottlenecks that require attention.'
    )

    doc.add_page_break()

    # ==================== PAGE 6 ====================
    doc.add_heading('6. Integrations', level=1)

    p_int1 = doc.add_paragraph()
    p_int1.add_run(
        'TaskFlow Pro connects with over 200 third-party services through its integration '
        'marketplace. Each integration is configured through a dedicated settings page where '
        'you authenticate with the external service and map data fields between systems. '
        'Popular integrations include Google Workspace, Dropbox, Salesforce, and Zendesk.'
    )

    p_int2 = doc.add_paragraph()
    p_int2.add_run(
        'The REST API provides programmatic access to all TaskFlow Pro features. API keys '
        'can be generated in the developer settings panel. The API supports CRUD operations '
        'on projects, tasks, comments, and attachments. Rate limiting is applied at 1000 '
        'requests per minute per API key to ensure platform stability.'
    )

    p_int3 = doc.add_paragraph()
    p_int3.add_run(
        'Webhooks enable real-time event-driven workflows. You can configure webhooks to fire '
        'on specific events such as task creation, status change, or comment addition. The '
        'webhook payload includes the full object data along with metadata about the triggering '
        'user and timestamp. Failed webhook deliveries are retried up to five times.'
    )

    doc.add_page_break()

    # ==================== PAGE 7 ====================
    doc.add_heading('7. Security and Permissions', level=1)

    p_sec1 = doc.add_paragraph()
    p_sec1.add_run(
        'Data security is a top priority for TaskFlow Pro. All data is encrypted at rest '
        'using AES-256 encryption and in transit using TLS 1.3. Two-factor authentication '
        'is available for all accounts and can be enforced at the organization level by '
        'administrators. Single sign-on via SAML 2.0 and OAuth 2.0 is supported.'
    )

    p_sec2 = doc.add_paragraph()
    p_sec2.add_run(
        'Role-based access control allows administrators to define granular permissions for '
        'each team member. Predefined roles include Viewer, Editor, Manager, and Admin. '
        'Custom roles can be created to accommodate unique organizational structures. '
        'Permissions can be set at the organization, project, or individual task level.'
    )

    p_sec3 = doc.add_paragraph()
    p_sec3.add_run(
        'Audit logs capture every action performed within the platform, including login events, '
        'permission changes, and data exports. Logs are retained for 12 months and can be '
        'downloaded by organization administrators for compliance purposes. Real-time alerts '
        'can be configured for suspicious activities such as unusual login patterns.'
    )

    doc.add_page_break()

    # ==================== PAGE 8 ====================
    doc.add_heading('8. Troubleshooting and Support', level=1)

    p_sup1 = doc.add_paragraph()
    p_sup1.add_run(
        'If you encounter issues while using TaskFlow Pro, the in-app help center provides '
        'searchable articles covering common questions and known issues. The help center is '
        'accessible from the question mark icon in the top navigation bar. Articles are '
        'organized by category and include step-by-step instructions with screenshots.'
    )

    p_sup2 = doc.add_paragraph()
    p_sup2.add_run(
        'For technical support, you can submit a ticket through the support portal or email '
        'support@taskflowpro.com directly. Premium plan customers have access to live chat '
        'support during business hours and receive priority ticket handling. Enterprise plan '
        'customers are assigned a dedicated account manager for personalized assistance.'
    )

    p_sup3 = doc.add_paragraph()
    p_sup3.add_run(
        'Community forums are available for peer-to-peer support and feature requests. The '
        'forums are moderated by TaskFlow Pro staff and experienced community members. Popular '
        'feature requests are reviewed quarterly and may be included in future product releases. '
        'You can vote on existing requests to help prioritize the development roadmap.'
    )

    # Footer
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p_footer.add_run('© 2026 TaskFlow Pro Inc. All rights reserved.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ==================== ADD TRACKED CHANGES ====================
    # Apply tracked changes by searching for text content, not by index
    tracked_changes = [
        (1, 'click', 'select', 'first project,'),
        (3, 'consolidated', 'comprehensive', 'provides a'),
        (5, 'manageable', 'actionable', 'into'),
        (7, 'simultaneously', 'concurrently', 'same task'),
        (9, 'powerful', 'robust', 'includes a'),
        (11, 'programmatic', 'direct', 'provides'),
        (13, 'granular', 'detailed', 'define'),
    ]

    paragraphs = doc.paragraphs
    for rev_id, old_text, new_text, context_hint in tracked_changes:
        found = False
        for para in paragraphs:
            if old_text in para.text and context_hint in para.text:
                result = add_tracked_change_replace(para._element, rev_id, old_text, new_text)
                print(f'TC (rev {rev_id}, {old_text}->{new_text}): {result}')
                found = True
                break
        if not found:
            # Fallback: search without context hint
            for para in paragraphs:
                if old_text in para.text:
                    result = add_tracked_change_replace(para._element, rev_id, old_text, new_text)
                    print(f'TC fallback (rev {rev_id}, {old_text}->{new_text}): {result}')
                    found = True
                    break
        if not found:
            print(f'TC FAILED (rev {rev_id}, {old_text}->{new_text}): not found')

    doc.save(OUTPUT)
    print(f'Document with tracked changes saved: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
