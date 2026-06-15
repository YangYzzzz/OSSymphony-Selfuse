"""
Initial Setup: Writer document with 4 sections and 8 endnotes (Roman numerals, end of document)
Task ID: writer_bs_018
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
import zipfile
import shutil
import tempfile

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_018'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

WML = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
MARKUP = 'http://schemas.openxmlformats.org/markup-compatibility/2006'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)

def create_initial():
    """
    Create a .docx with 4 sections and 8 endnotes (2 per section).
    Endnotes use Roman numeral formatting (i, ii, iii...) and appear at end of document.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.ns import qn
    from lxml import etree

    doc = Document()

    # -- Section 1: Introduction to Renewable Energy --
    h1 = doc.add_heading('Introduction to Renewable Energy', level=1)
    p1a = doc.add_paragraph(
        'The global transition toward renewable energy sources has accelerated '
        'dramatically since 2020. Solar photovoltaic installations grew by 35% '
        'in 2024 alone, driven by declining manufacturing costs and supportive '
        'government policies across major economies.'
    )
    p1b = doc.add_paragraph(
        'Wind energy capacity additions reached 120 GW globally in 2024, '
        'with offshore wind projects accounting for nearly a quarter of new '
        'installations. The levelized cost of energy from onshore wind has '
        'dropped below $30 per megawatt-hour in several markets.'
    )

    # -- Section 2: Economic Impact Assessment --
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    h2 = doc.add_heading('Economic Impact Assessment', level=1)
    p2a = doc.add_paragraph(
        'The renewable energy sector employed approximately 14.5 million people '
        'worldwide in 2024, representing a 12% increase from the previous year. '
        'The solar industry alone accounted for 5.2 million jobs, concentrated '
        'primarily in manufacturing, installation, and maintenance roles.'
    )
    p2b = doc.add_paragraph(
        'Investment in clean energy technologies reached $1.8 trillion in 2024, '
        'surpassing fossil fuel investment for the first time. Private equity '
        'and venture capital funding for energy storage solutions grew by 45%, '
        'reflecting investor confidence in next-generation battery technologies.'
    )

    # -- Section 3: Environmental Benefits --
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    h3 = doc.add_heading('Environmental Benefits', level=1)
    p3a = doc.add_paragraph(
        'Carbon dioxide emissions from the power sector declined by 8% in 2024 '
        'compared to 2023, largely attributable to the displacement of coal-fired '
        'generation by renewable sources. The European Union achieved a 22% '
        'reduction in power-sector emissions relative to its 2019 baseline.'
    )
    p3b = doc.add_paragraph(
        'Water consumption in electricity generation decreased by 15% as thermal '
        'power plants were retired in favor of solar and wind installations, '
        'which require minimal water during operation. This shift has been '
        'particularly significant in water-stressed regions of South Asia and Africa.'
    )

    # -- Section 4: Future Outlook and Challenges --
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    h4 = doc.add_heading('Future Outlook and Challenges', level=1)
    p4a = doc.add_paragraph(
        'Grid integration remains the primary technical challenge for renewable '
        'energy deployment. As variable generation sources exceed 40% of total '
        'capacity in leading markets, investments in transmission infrastructure '
        'and demand-side flexibility become critical enablers of further growth.'
    )
    p4b = doc.add_paragraph(
        'The International Energy Agency projects that renewable sources will '
        'account for 50% of global electricity generation by 2030, contingent '
        'upon continued policy support and resolution of supply chain constraints '
        'for critical minerals including lithium, cobalt, and rare earth elements.'
    )

    # Save the basic document first
    doc.save(OUTPUT)

    # Now modify the raw XML to add endnotes with Roman numeral formatting
    # python-docx doesn't support endnotes natively, so we use lxml + zipfile

    tmpdir = tempfile.mkdtemp()
    try:
        # Extract the docx
        with zipfile.ZipFile(OUTPUT, 'r') as z:
            z.extractall(tmpdir)

        # --- Create endnotes.xml ---
        endnotes_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:endnotes xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
            xmlns:mc="{MARKUP}"
            xmlns:o="urn:schemas-microsoft-com:office:office"
            xmlns:r="{REL}"
            xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            xmlns:w10="urn:schemas-microsoft-com:office:word"
            xmlns:w="{WML}"
            xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml">
  <w:endnote w:type="separator" w:id="-1">
    <w:p>
      <w:r>
        <w:separator/>
      </w:r>
    </w:p>
  </w:endnote>
  <w:endnote w:type="continuationSeparator" w:id="0">
    <w:p>
      <w:r>
        <w:continuationSeparator/>
      </w:r>
    </w:p>
  </w:endnote>
  <w:endnote w:id="1">
    <w:p>
      <w:pPr><w:pStyle w:val="EndnoteText"/></w:pPr>
      <w:r><w:rPr><w:rStyle w:val="EndnoteReference"/></w:rPr><w:endnoteRef/></w:r>
      <w:r><w:t xml:space="preserve"> International Renewable Energy Agency, "Renewable Capacity Statistics 2025," IRENA Report, Abu Dhabi, 2025.</w:t></w:r>
    </w:p>
  </w:endnote>
  <w:endnote w:id="2">
    <w:p>
      <w:pPr><w:pStyle w:val="EndnoteText"/></w:pPr>
      <w:r><w:rPr><w:rStyle w:val="EndnoteReference"/></w:rPr><w:endnoteRef/></w:r>
      <w:r><w:t xml:space="preserve"> Global Wind Energy Council, "Global Wind Report 2025," Brussels, 2025.</w:t></w:r>
    </w:p>
  </w:endnote>
  <w:endnote w:id="3">
    <w:p>
      <w:pPr><w:pStyle w:val="EndnoteText"/></w:pPr>
      <w:r><w:rPr><w:rStyle w:val="EndnoteReference"/></w:rPr><w:endnoteRef/></w:r>
      <w:r><w:t xml:space="preserve"> International Labour Organization, "World Employment and Social Outlook: Greening with Jobs," Geneva, 2024.</w:t></w:r>
    </w:p>
  </w:endnote>
  <w:endnote w:id="4">
    <w:p>
      <w:pPr><w:pStyle w:val="EndnoteText"/></w:pPr>
      <w:r><w:rPr><w:rStyle w:val="EndnoteReference"/></w:rPr><w:endnoteRef/></w:r>
      <w:r><w:t xml:space="preserve"> BloombergNEF, "Energy Transition Investment Trends 2025," Bloomberg L.P., New York, 2025.</w:t></w:r>
    </w:p>
  </w:endnote>
  <w:endnote w:id="5">
    <w:p>
      <w:pPr><w:pStyle w:val="EndnoteText"/></w:pPr>
      <w:r><w:rPr><w:rStyle w:val="EndnoteReference"/></w:rPr><w:endnoteRef/></w:r>
      <w:r><w:t xml:space="preserve"> European Environment Agency, "Trends and Projections in Europe 2025," Copenhagen, 2025.</w:t></w:r>
    </w:p>
  </w:endnote>
  <w:endnote w:id="6">
    <w:p>
      <w:pPr><w:pStyle w:val="EndnoteText"/></w:pPr>
      <w:r><w:rPr><w:rStyle w:val="EndnoteReference"/></w:rPr><w:endnoteRef/></w:r>
      <w:r><w:t xml:space="preserve"> World Resources Institute, "Water Stress and Energy: A Global Assessment," Washington D.C., 2024.</w:t></w:r>
    </w:p>
  </w:endnote>
  <w:endnote w:id="7">
    <w:p>
      <w:pPr><w:pStyle w:val="EndnoteText"/></w:pPr>
      <w:r><w:rPr><w:rStyle w:val="EndnoteReference"/></w:rPr><w:endnoteRef/></w:r>
      <w:r><w:t xml:space="preserve"> MIT Energy Initiative, "The Future of Grid Integration," Cambridge, MA, 2025.</w:t></w:r>
    </w:p>
  </w:endnote>
  <w:endnote w:id="8">
    <w:p>
      <w:pPr><w:pStyle w:val="EndnoteText"/></w:pPr>
      <w:r><w:rPr><w:rStyle w:val="EndnoteReference"/></w:rPr><w:endnoteRef/></w:r>
      <w:r><w:t xml:space="preserve"> International Energy Agency, "World Energy Outlook 2025," Paris, 2025.</w:t></w:r>
    </w:p>
  </w:endnote>
</w:endnotes>'''

        endnotes_path = os.path.join(tmpdir, 'word', 'endnotes.xml')
        with open(endnotes_path, 'w', encoding='utf-8') as f:
            f.write(endnotes_xml)

        # --- Modify document.xml to add endnote references ---
        doc_path = os.path.join(tmpdir, 'word', 'document.xml')
        tree = etree.parse(doc_path)
        root = tree.getroot()
        nsmap = {'w': WML}

        body = root.find('.//w:body', nsmap)
        paragraphs = body.findall('.//w:p', nsmap)

        # We need to add endnote references to specific paragraphs.
        # The structure: each section has content paragraphs, we add endnote refs to them.
        # Map: paragraph indices in body -> endnote IDs
        # Body structure: heading, para, para, sectPr(in last para pPr), heading, para, para, ...
        # Let's find all paragraphs and add refs to the content ones.

        # Collect all direct-child paragraphs of body
        body_children = list(body)
        all_paras = [c for c in body_children if c.tag == f'{{{WML}}}p']

        # We have: h1, p1a, p1b, | h2, p2a, p2b, | h3, p3a, p3b, | h4, p4a, p4b
        # indices:  0,   1,   2,     3,   4,   5,     6,   7,   8,     9,  10,  11
        # Endnote assignments: p1a->1, p1b->2, p2a->3, p2b->4, p3a->5, p3b->6, p4a->7, p4b->8
        endnote_assignments = {
            1: 1, 2: 2,    # Section 1 paragraphs
            4: 3, 5: 4,    # Section 2 paragraphs
            7: 5, 8: 6,    # Section 3 paragraphs
            10: 7, 11: 8,  # Section 4 paragraphs
        }

        for para_idx, endnote_id in endnote_assignments.items():
            if para_idx < len(all_paras):
                para = all_paras[para_idx]
                # Add endnote reference run at the end of the paragraph
                run_elem = etree.SubElement(para, f'{{{WML}}}r')
                rpr = etree.SubElement(run_elem, f'{{{WML}}}rPr')
                rstyle = etree.SubElement(rpr, f'{{{WML}}}rStyle')
                rstyle.set(f'{{{WML}}}val', 'EndnoteReference')
                endnote_ref = etree.SubElement(run_elem, f'{{{WML}}}endnoteReference')
                endnote_ref.set(f'{{{WML}}}id', str(endnote_id))

        # --- Set endnote numbering format to Roman numerals (lowerRoman) at document level ---
        # Add endnotePr to the last sectPr in body (document-level settings)
        sect_prs = body.findall('.//w:sectPr', nsmap)
        for sect_pr in sect_prs:
            # Add endnotePr with numFmt=lowerRoman and pos=docEnd (default)
            endnote_pr = etree.SubElement(sect_pr, f'{{{WML}}}endnotePr')
            num_fmt = etree.SubElement(endnote_pr, f'{{{WML}}}numFmt')
            num_fmt.set(f'{{{WML}}}val', 'lowerRoman')

        # Save modified document.xml
        tree.write(doc_path, xml_declaration=True, encoding='UTF-8', standalone=True)

        # --- Add endnotes relationship to word/_rels/document.xml.rels ---
        rels_path = os.path.join(tmpdir, 'word', '_rels', 'document.xml.rels')
        rels_tree = etree.parse(rels_path)
        rels_root = rels_tree.getroot()
        rels_ns = 'http://schemas.openxmlformats.org/package/2006/relationships'

        # Find max rId
        max_id = 0
        for rel in rels_root:
            rid = rel.get('Id', '')
            if rid.startswith('rId'):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass

        new_rel = etree.SubElement(rels_root, f'{{{rels_ns}}}Relationship')
        new_rel.set('Id', f'rId{max_id + 1}')
        new_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes')
        new_rel.set('Target', 'endnotes.xml')

        rels_tree.write(rels_path, xml_declaration=True, encoding='UTF-8', standalone=True)

        # --- Update [Content_Types].xml to include endnotes.xml ---
        ct_path = os.path.join(tmpdir, '[Content_Types].xml')
        ct_tree = etree.parse(ct_path)
        ct_root = ct_tree.getroot()
        ct_ns = 'http://schemas.openxmlformats.org/package/2006/content-types'

        override = etree.SubElement(ct_root, f'{{{ct_ns}}}Override')
        override.set('PartName', '/word/endnotes.xml')
        override.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml')

        ct_tree.write(ct_path, xml_declaration=True, encoding='UTF-8', standalone=True)

        # --- Add EndnoteText and EndnoteReference styles if not present ---
        styles_path = os.path.join(tmpdir, 'word', 'styles.xml')
        styles_tree = etree.parse(styles_path)
        styles_root = styles_tree.getroot()

        # Check if EndnoteText style exists
        existing_styles = [s.get(f'{{{WML}}}styleId') for s in styles_root.findall(f'{{{WML}}}style')]

        if 'EndnoteText' not in existing_styles:
            style = etree.SubElement(styles_root, f'{{{WML}}}style')
            style.set(f'{{{WML}}}type', 'paragraph')
            style.set(f'{{{WML}}}styleId', 'EndnoteText')
            name = etree.SubElement(style, f'{{{WML}}}name')
            name.set(f'{{{WML}}}val', 'endnote text')
            rpr = etree.SubElement(style, f'{{{WML}}}rPr')
            sz = etree.SubElement(rpr, f'{{{WML}}}sz')
            sz.set(f'{{{WML}}}val', '20')  # 10pt

        if 'EndnoteReference' not in existing_styles:
            style = etree.SubElement(styles_root, f'{{{WML}}}style')
            style.set(f'{{{WML}}}type', 'character')
            style.set(f'{{{WML}}}styleId', 'EndnoteReference')
            name = etree.SubElement(style, f'{{{WML}}}name')
            name.set(f'{{{WML}}}val', 'endnote reference')
            rpr = etree.SubElement(style, f'{{{WML}}}rPr')
            vert = etree.SubElement(rpr, f'{{{WML}}}vertAlign')
            vert.set(f'{{{WML}}}val', 'superscript')

        styles_tree.write(styles_path, xml_declaration=True, encoding='UTF-8', standalone=True)

        # --- Repack the docx ---
        os.remove(OUTPUT)
        with zipfile.ZipFile(OUTPUT, 'w', zipfile.ZIP_DEFLATED) as zout:
            for dirpath, dirnames, filenames in os.walk(tmpdir):
                for fn in filenames:
                    full = os.path.join(dirpath, fn)
                    arcname = os.path.relpath(full, tmpdir)
                    zout.write(full, arcname)

        print(f'Initial file created: {OUTPUT}')

    finally:
        shutil.rmtree(tmpdir)

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')

create_initial()
