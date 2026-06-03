"""
Initial Setup: Vendor Agreement with Find & Replace task
Task ID: writer_legal_035
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_035'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    if alignment:
        p.paragraph_format.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ---- TITLE ----
    # Vendor count: 1 (title)
    add_para(doc, 'VENDOR SERVICES AGREEMENT', bold=True,
             alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=Pt(12))

    add_para(doc, 'Effective Date: March 15, 2025',
             alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_para(doc, 'Contract Reference: VSA-2025-0471',
             alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=Pt(18))

    # ---- PREAMBLE ----  Vendor count: 5
    add_para(doc,
        'This Vendor Services Agreement ("Agreement") is entered into by and between '
        'Meridian Technologies Inc., a Delaware corporation with its principal offices at '
        '2200 Innovation Drive, Suite 400, San Jose, CA 95134 ("Client"), and '
        'Pinnacle Solutions Group LLC, a California limited liability company with its '
        'principal offices at 780 Market Street, Floor 12, San Francisco, CA 94102 '
        '("Vendor").')

    add_para(doc,
        'WHEREAS, the Client desires to engage the Vendor for certain professional '
        'services as described herein; and WHEREAS, the Vendor represents that it has '
        'the expertise, resources, and capacity to perform such services;')

    add_para(doc,
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements '
        'set forth herein, and for other good and valuable consideration, the receipt '
        'and sufficiency of which are hereby acknowledged, the parties agree as follows:')

    # ---- SECTION 1: DEFINITIONS ----  Vendor count: 5
    add_heading(doc, '1. Definitions', level=1)

    add_para(doc,
        '1.1 "Vendor Personnel" means any employees, agents, contractors, or '
        'subcontractors engaged by the Vendor to perform Services under this Agreement.')

    add_para(doc,
        '1.2 "Vendor Materials" means any tools, methodologies, software, documentation, '
        'or other materials owned by or licensed to the Vendor prior to the Effective Date.')

    add_para(doc,
        '1.3 "Services" means the professional, technical, and consulting services '
        'to be provided by the Vendor as described in Exhibit A attached hereto.')

    add_para(doc,
        '1.4 "Deliverables" means all work products, reports, software, documentation, '
        'and other tangible outputs created by the Vendor in connection with the Services.')

    # ---- SECTION 2: SCOPE OF SERVICES ----  Vendor count: 6
    add_heading(doc, '2. Scope of Services', level=1)

    add_para(doc,
        '2.1 The Vendor shall provide the Services described in Exhibit A in accordance '
        'with the timelines, specifications, and standards set forth therein. Qualified '
        'Vendor Personnel shall be assigned to perform all Services.')

    add_para(doc,
        '2.2 The Vendor shall not subcontract or delegate any portion of the Services '
        'without the prior written consent of the Client. Any approved subcontractor '
        'shall be bound by the same obligations as the Vendor under this Agreement.')

    add_para(doc,
        '2.3 The Vendor agrees to perform all Services in a professional and workmanlike '
        'manner, consistent with industry standards and best practices applicable to '
        'the type of services being provided.')

    # ---- SECTION 3: COMPENSATION ----  Vendor count: 5
    add_heading(doc, '3. Compensation and Payment', level=1)

    add_para(doc,
        '3.1 In consideration for the Services, the Client shall pay the Vendor the fees '
        'set forth in Exhibit B ("Fee Schedule"). Unless otherwise specified, the Vendor '
        'shall invoice the Client on a monthly basis.')

    add_para(doc,
        '3.2 The Client shall pay each undisputed Vendor invoice within thirty (30) '
        'calendar days of receipt. The Vendor shall provide sufficient documentation '
        'to support all invoiced amounts, including timesheets and expense reports.')

    add_para(doc,
        '3.3 The Vendor shall be responsible for all taxes, insurance, and other '
        'statutory obligations arising from payments made under this Agreement. The '
        'Vendor acknowledges that it is an independent contractor.')

    # ---- SECTION 4: TERM AND TERMINATION ----  Vendor count: 5
    add_heading(doc, '4. Term and Termination', level=1)

    add_para(doc,
        '4.1 This Agreement shall commence on the Effective Date and continue for '
        'a period of twenty-four (24) months unless earlier terminated. Either party '
        'may terminate this Agreement with sixty (60) days written notice.')

    add_para(doc,
        '4.2 The Client may terminate this Agreement immediately if the Vendor '
        'materially breaches any provision and fails to cure such breach within '
        'fifteen (15) business days after receiving written notice of the breach.')

    add_para(doc,
        '4.3 Upon termination, the Vendor shall promptly return all Client property, '
        'Confidential Information, and incomplete Deliverables. The Vendor shall '
        'cooperate fully with the transition of Services to the Client or a successor vendor.')

    # ---- SECTION 5: CONFIDENTIALITY ----  Vendor count: 4
    add_heading(doc, '5. Confidentiality', level=1)

    add_para(doc,
        '5.1 The Vendor acknowledges that in the course of performing Services, '
        'the Vendor may receive or have access to confidential and proprietary '
        'information of the Client ("Confidential Information").')

    add_para(doc,
        '5.2 The Vendor shall hold all Confidential Information in strict confidence '
        'and shall not disclose such information to any third party without the '
        'prior written consent of the Client.')

    add_para(doc,
        '5.3 The obligations of confidentiality shall survive the termination '
        'of this Agreement for a period of five (5) years. The Vendor shall ensure '
        'that all Vendor Personnel are bound by equivalent confidentiality obligations.')

    # ---- SECTION 6: INTELLECTUAL PROPERTY ----  Vendor count: 4
    add_heading(doc, '6. Intellectual Property', level=1)

    add_para(doc,
        '6.1 All Deliverables and work product created by the Vendor in the performance '
        'of Services shall be the sole and exclusive property of the Client. The Vendor '
        'hereby assigns all rights, title, and interest in such work product to the Client.')

    add_para(doc,
        '6.2 The Vendor retains ownership of all Vendor Materials. To the extent any '
        'Vendor Materials are incorporated into Deliverables, the Vendor grants the '
        'Client a perpetual, non-exclusive, royalty-free license to use such materials.')

    # ---- SECTION 7: WARRANTIES AND REPRESENTATIONS ----  Vendor count: 4
    add_heading(doc, '7. Warranties and Representations', level=1)

    add_para(doc,
        '7.1 The Vendor represents and warrants that: (a) it has the authority to '
        'enter into this Agreement; (b) the Services will be performed in a competent '
        'and professional manner; and (c) the Deliverables will conform to the '
        'specifications set forth in Exhibit A.')

    add_para(doc,
        '7.2 The Vendor further warrants that no Deliverable will infringe upon '
        'or misappropriate any intellectual property rights of any third party. '
        'The Vendor shall indemnify and hold harmless the Client from any claims '
        'arising from such infringement.')

    # ---- SECTION 8: LIABILITY AND INDEMNIFICATION ----  Vendor count: 3
    add_heading(doc, '8. Liability and Indemnification', level=1)

    add_para(doc,
        '8.1 The Vendor shall indemnify, defend, and hold harmless the Client and '
        'its officers, directors, employees, and agents from and against any and all '
        'claims, damages, losses, and expenses arising out of the Vendor\'s performance '
        'or failure to perform under this Agreement.')

    add_para(doc,
        '8.2 In no event shall either party be liable for any indirect, incidental, '
        'consequential, or punitive damages. The Vendor\'s total aggregate liability '
        'shall not exceed the total fees paid under this Agreement during the preceding '
        'twelve (12) month period.')

    # ---- SECTION 9: GENERAL PROVISIONS ----  Vendor count: 3
    add_heading(doc, '9. General Provisions', level=1)

    add_para(doc,
        '9.1 This Agreement constitutes the entire agreement between the Client and '
        'the Vendor with respect to the subject matter hereof. No modification of '
        'this Agreement shall be valid unless made in writing and signed by both parties.')

    add_para(doc,
        '9.2 The Vendor shall comply with all applicable federal, state, and local '
        'laws, regulations, and ordinances in the performance of Services under this '
        'Agreement.')

    add_para(doc,
        '9.3 This Agreement shall be governed by and construed in accordance with '
        'the laws of the State of California. Any dispute arising under this Agreement '
        'shall be resolved by binding arbitration in San Francisco, California.')

    # ---- SIGNATURE BLOCK ----
    add_para(doc, '', space_after=Pt(24))

    add_para(doc, 'IN WITNESS WHEREOF, the parties have executed this Agreement '
        'as of the date first written above.', space_after=Pt(18))

    add_para(doc, 'CLIENT: Meridian Technologies Inc.', bold=True)
    add_para(doc, 'By: _________________________')
    add_para(doc, 'Name: Jonathan R. Whitfield')
    add_para(doc, 'Title: Chief Operating Officer')
    add_para(doc, 'Date: March 15, 2025', space_after=Pt(18))

    # Vendor count in signature: 0 (just the word "VENDOR" as a label — uppercase is fine,
    # but per task it should be counted)
    # Actually we need to be precise. Let me keep this as a label.
    # This does NOT count as 'Vendor' since it's all caps 'VENDOR'.
    add_para(doc, 'VENDOR: Pinnacle Solutions Group LLC', bold=True)
    add_para(doc, 'By: _________________________')
    add_para(doc, 'Name: Rebecca A. Torres')
    add_para(doc, 'Title: Managing Director')
    add_para(doc, 'Date: March 15, 2025', space_after=Pt(18))

    # ---- CONTACT INFORMATION (contains lowercase 'vendor') ----
    add_heading(doc, 'Contact Information', level=2)

    add_para(doc,
        'Client Contact: Jonathan Whitfield, jwhitfield@meridiantech.com, '
        '(408) 555-2200')
    # lowercase 'vendor' in URL/email — these should NOT be replaced
    add_para(doc,
        'Vendor Contact: Rebecca Torres, rtorres@pinnaclesolutions.com, '
        '(415) 555-7800')
    add_para(doc,
        'Vendor Portal: https://portal.pinnaclesolutions.com/vendor-services/login')
    add_para(doc,
        'Support: support@pinnaclesolutions.com/vendor-helpdesk')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Count occurrences for verification
    doc2 = Document(OUTPUT)
    vendor_cap = 0
    vendor_low = 0
    full_text = ''
    for para in doc2.paragraphs:
        full_text += para.text + '\n'

    import re
    vendor_cap = len(re.findall(r'Vendor', full_text))
    vendor_low = len(re.findall(r'vendor', full_text))
    print(f'Capitalized "Vendor" count: {vendor_cap}')
    print(f'Lowercase "vendor" count: {vendor_low}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
