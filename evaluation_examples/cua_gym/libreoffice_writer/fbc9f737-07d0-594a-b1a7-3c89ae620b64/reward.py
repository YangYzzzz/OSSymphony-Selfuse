"""
Reward Script: Case-sensitive Find & Replace — 'Vendor' -> 'Service Provider'
Task ID: writer_legal_035
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): All 'Vendor' (capital V) instances removed (0 remaining)
  Component 2 (0.3): 'Service Provider' appears at least 45 times
  Component 3 (0.3): Case-sensitivity verified — lowercase 'vendor' preserved AND 'Vendor' removed
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_035'


def count_pattern_in_doc(doc, pattern):
    """Count regex pattern occurrences across all paragraphs and table cells."""
    count = 0
    for para in doc.paragraphs:
        count += len(re.findall(pattern, para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count += len(re.findall(pattern, para.text))
    return count


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Count occurrences
    vendor_upper_count = count_pattern_in_doc(doc, r'Vendor')
    service_provider_count = count_pattern_in_doc(doc, r'Service Provider')
    # lowercase-only 'vendor' — not preceded by uppercase letter
    vendor_lower_count = count_pattern_in_doc(doc, r'\bvendor\b')

    print(f"Found: 'Vendor' (uppercase V) = {vendor_upper_count}")
    print(f"Found: 'Service Provider' = {service_provider_count}")
    print(f"Found: 'vendor' (lowercase) = {vendor_lower_count}")

    # Component 1: All 'Vendor' (capital V) instances removed (0.4 points)
    # Initial has 45. After task, should be 0.
    try:
        if vendor_upper_count == 0:
            print(f"PASS: Component 1 — No 'Vendor' instances remain (0.4 pts)")
            total_score += 0.4
        else:
            # Partial credit: fraction that were removed (45 - remaining) / 45
            removed = max(0, 45 - vendor_upper_count)
            if removed > 0:
                partial = round(0.4 * (removed / 45), 2)
                print(f"PARTIAL: Component 1 — {vendor_upper_count} 'Vendor' still present, {removed}/45 removed ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 1 — All 45 'Vendor' instances still present")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 'Service Provider' appears at least 45 times (0.3 points)
    # Initial has 0. After task, should be 45.
    try:
        if service_provider_count >= 45:
            print(f"PASS: Component 2 — 'Service Provider' count = {service_provider_count} (>= 45) (0.3 pts)")
            total_score += 0.3
        elif service_provider_count > 0:
            partial = round(0.3 * (service_provider_count / 45), 2)
            print(f"PARTIAL: Component 2 — 'Service Provider' count = {service_provider_count}/45 ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — 'Service Provider' count = 0 (expected >= 45)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Case-sensitivity correctly applied (0.3 points)
    # Compound check: lowercase 'vendor' must still exist (>= 3) AND 'Vendor' must be gone (== 0)
    # This ensures the replacement was case-sensitive. On initial_env, 'Vendor' is still 45
    # so this compound check fails even though lowercase vendor exists.
    try:
        if vendor_lower_count >= 3 and vendor_upper_count == 0:
            print(f"PASS: Component 3 — Case-sensitive replacement verified: {vendor_lower_count} lowercase 'vendor' preserved, 0 uppercase 'Vendor' (0.3 pts)")
            total_score += 0.3
        elif vendor_lower_count >= 3 and vendor_upper_count > 0:
            print(f"FAIL: Component 3 — Lowercase 'vendor' preserved ({vendor_lower_count}) but uppercase 'Vendor' not fully removed ({vendor_upper_count} remain)")
        elif vendor_lower_count == 0 and vendor_upper_count == 0:
            print(f"FAIL: Component 3 — Lowercase 'vendor' was also replaced ({vendor_lower_count} found) — replacement was NOT case-sensitive")
        else:
            print(f"FAIL: Component 3 — vendor_lower={vendor_lower_count}, Vendor_upper={vendor_upper_count}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice state before verification
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
