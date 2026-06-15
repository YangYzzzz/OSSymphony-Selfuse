"""
Initial Setup: Legal document with heading levels but no outline numbering
Task ID: writer_bs_081
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_081'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # -- Title --
    title = doc.add_heading('Master Services Agreement', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        'This Master Services Agreement ("Agreement") is entered into as of '
        'March 15, 2025, by and between Greenfield Technologies, Inc., a Delaware '
        'corporation ("Provider"), and Meridian Health Systems, LLC, an Illinois '
        'limited liability company ("Client").'
    )
    intro.paragraph_format.space_after = Pt(12)

    # =========================================================
    # HEADING LEVEL 1 — "Definitions" (plain, no numbering)
    # =========================================================
    doc.add_heading('Definitions', level=1)

    # Level 2 headings under Definitions
    doc.add_heading('General Terms', level=2)

    doc.add_paragraph(
        '"Confidential Information" means any non-public information disclosed by '
        'either party to the other party, whether orally, in writing, or by inspection '
        'of tangible objects, that is designated as confidential or that reasonably '
        'should be understood to be confidential.'
    )

    doc.add_paragraph(
        '"Deliverables" means all work products, reports, analyses, software code, '
        'documentation, and other materials developed by Provider in the performance '
        'of Services under this Agreement.'
    )

    # Level 3 headings
    doc.add_heading('First provision', level=3)
    doc.add_paragraph(
        'The term "Effective Date" shall mean the date first written above, unless '
        'otherwise specified in a Statement of Work executed by both parties.'
    )

    doc.add_heading('Second provision', level=3)
    doc.add_paragraph(
        '"Intellectual Property Rights" means all patents, copyrights, trademarks, '
        'trade secrets, and any other proprietary rights recognized under applicable law.'
    )

    doc.add_heading('Third provision', level=3)
    doc.add_paragraph(
        '"Service Level Agreement" or "SLA" means the performance standards and '
        'metrics set forth in Exhibit B attached hereto.'
    )

    doc.add_heading('Technical Terms', level=2)

    doc.add_paragraph(
        '"API" means Application Programming Interface, including RESTful and '
        'GraphQL interfaces provided by the Platform.'
    )

    doc.add_paragraph(
        '"Platform" means the cloud-based software solution operated by Provider, '
        'including all associated databases, servers, and network infrastructure.'
    )

    doc.add_heading('First provision', level=3)
    doc.add_paragraph(
        '"Uptime" shall be calculated as the total number of minutes in a calendar '
        'month minus the number of minutes of Downtime, divided by the total number '
        'of minutes in that month, expressed as a percentage.'
    )

    doc.add_heading('Second provision', level=3)
    doc.add_paragraph(
        '"Data Processing" means any operation performed on personal data, including '
        'collection, recording, organization, structuring, storage, adaptation, '
        'retrieval, consultation, use, disclosure, or erasure.'
    )

    # =========================================================
    # HEADING LEVEL 1 — "Scope of Services"
    # =========================================================
    doc.add_heading('Scope of Services', level=1)

    doc.add_heading('Service Description', level=2)

    doc.add_paragraph(
        'Provider shall deliver the following services to Client in accordance '
        'with the terms and conditions of this Agreement and the applicable '
        'Statement of Work:'
    )

    doc.add_heading('Implementation services', level=3)
    doc.add_paragraph(
        'Provider shall install, configure, and deploy the Platform in Client\'s '
        'production environment within sixty (60) calendar days of the Effective Date. '
        'Implementation includes data migration from Client\'s legacy systems, '
        'user acceptance testing, and production go-live support.'
    )

    doc.add_heading('Training and support', level=3)
    doc.add_paragraph(
        'Provider shall conduct up to forty (40) hours of on-site training for '
        'Client\'s designated personnel, covering system administration, end-user '
        'operations, and reporting functionality.'
    )

    doc.add_heading('Service Modifications', level=2)

    doc.add_paragraph(
        'Either party may request modifications to the Services by submitting a '
        'written change request. All change requests shall be reviewed jointly, '
        'and no modification shall be effective until both parties execute an '
        'amended Statement of Work.'
    )

    doc.add_heading('Approval process', level=3)
    doc.add_paragraph(
        'Change requests must be submitted no fewer than fifteen (15) business '
        'days prior to the proposed implementation date, and shall include a '
        'detailed description of the requested changes, estimated costs, and '
        'projected timeline.'
    )

    # =========================================================
    # HEADING LEVEL 1 — "Compensation and Payment"
    # =========================================================
    doc.add_heading('Compensation and Payment', level=1)

    doc.add_heading('Fee Structure', level=2)

    doc.add_paragraph(
        'Client shall pay Provider the fees set forth in Exhibit A, which may '
        'include a combination of fixed fees, time-and-materials charges, and '
        'recurring subscription fees. The base annual subscription fee shall be '
        'Four Hundred Fifty Thousand Dollars ($450,000), payable in equal '
        'quarterly installments of One Hundred Twelve Thousand Five Hundred '
        'Dollars ($112,500).'
    )

    doc.add_heading('Invoicing terms', level=3)
    doc.add_paragraph(
        'Provider shall submit invoices to Client on a monthly basis, no later '
        'than the fifth (5th) business day following the end of each calendar month. '
        'Each invoice shall include a detailed breakdown of services rendered, '
        'hours expended, and expenses incurred.'
    )

    doc.add_heading('Late payment', level=3)
    doc.add_paragraph(
        'Any amounts not paid within thirty (30) days of the invoice date shall '
        'accrue interest at the rate of one and one-half percent (1.5%) per month, '
        'or the maximum rate permitted by applicable law, whichever is less.'
    )

    doc.add_heading('Payment Methods', level=2)

    doc.add_paragraph(
        'All payments shall be made by wire transfer to Provider\'s designated '
        'bank account or via ACH transfer. Client shall be responsible for any '
        'bank fees or transfer charges associated with payment.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
