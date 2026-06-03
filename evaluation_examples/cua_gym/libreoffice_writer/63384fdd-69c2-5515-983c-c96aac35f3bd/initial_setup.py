"""
Initial Setup: Alert document with rectangle shape (blue fill, black border, no text)
Task ID: writer_obj_078
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import lxml.etree as etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'alert_doc'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Header ---
    heading = doc.add_heading('Safety & Compliance Notice', level=1)
    heading.paragraph_format.space_after = Pt(12)

    # --- Intro paragraph ---
    intro = doc.add_paragraph(
        'This document outlines critical safety protocols and compliance requirements '
        'for all staff members operating within the facility. Please read all sections '
        'carefully and acknowledge receipt by signing the attached form.'
    )
    intro.paragraph_format.space_after = Pt(8)

    # --- Rectangle shape (10cm x 3cm), blue fill (#0000FF), thin black border, no text ---
    # In DrawingML: 10cm = 3600000 EMU, 3cm = 1080000 EMU
    # Blue fill: 0000FF, thin black border: 1pt = 12700 EMU
    shape_para = doc.add_paragraph()
    shape_para.paragraph_format.space_before = Pt(6)
    shape_para.paragraph_format.space_after = Pt(6)
    shape_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    shape_run = shape_para.add_run()

    # DrawingML rectangle: blue solid fill, thin (1pt) black border, no text body
    drawing_xml = (
        '<w:drawing '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wp:inline distT="0" distB="0" distL="0" distR="0">'
        '<wp:extent cx="3600000" cy="1080000"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:docPr id="1" name="Rectangle 1" descr="Alert rectangle shape"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp>'
        '<wps:cNvSpPr>'
        '<a:spLocks noChangeArrowheads="1"/>'
        '</wps:cNvSpPr>'
        '<wps:spPr>'
        '<a:xfrm>'
        '<a:off x="0" y="0"/>'
        '<a:ext cx="3600000" cy="1080000"/>'
        '</a:xfrm>'
        '<a:prstGeom prst="rect">'
        '<a:avLst/>'
        '</a:prstGeom>'
        '<a:solidFill>'
        '<a:srgbClr val="0000FF"/>'
        '</a:solidFill>'
        '<a:ln w="12700">'
        '<a:solidFill>'
        '<a:srgbClr val="000000"/>'
        '</a:solidFill>'
        '</a:ln>'
        '</wps:spPr>'
        '<wps:bodyPr/>'
        '</wps:wsp>'
        '</a:graphicData>'
        '</a:graphic>'
        '</wp:inline>'
        '</w:drawing>'
    )

    drawing_elem = etree.fromstring(drawing_xml)
    shape_run._element.append(drawing_elem)

    # --- Section 1: General Safety Guidelines ---
    doc.add_heading('1. General Safety Guidelines', level=2)
    doc.add_paragraph(
        'All personnel must wear appropriate personal protective equipment (PPE) when '
        'entering designated hazard zones. Equipment includes hard hats, safety vests, '
        'steel-toed boots, and eye protection where required.'
    )

    # --- Section 2: Emergency Procedures ---
    doc.add_heading('2. Emergency Procedures', level=2)
    doc.add_paragraph(
        'In the event of an emergency, all staff must proceed to the nearest emergency '
        'exit and gather at the designated muster point. Do not use elevators during '
        'fire alarms or power outages.'
    )

    # Emergency contact list
    doc.add_paragraph('Emergency Contacts:', style='List Bullet')
    doc.add_paragraph('Fire Department: 119', style='List Bullet')
    doc.add_paragraph('Medical Emergency: 120', style='List Bullet')
    doc.add_paragraph('Security Office: ext. 5500', style='List Bullet')

    # --- Section 3: Compliance Requirements ---
    doc.add_heading('3. Compliance Requirements', level=2)
    doc.add_paragraph(
        'All employees are required to complete the annual safety training by March 31st '
        'each year. Failure to complete training may result in restricted access to '
        'certain operational areas pending certification renewal.'
    )

    # --- Footer note ---
    last_para = doc.add_paragraph(
        'Document last updated: February 2026 | Version 3.2 | Compliance Department'
    )
    last_para.paragraph_format.space_before = Pt(12)

    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
