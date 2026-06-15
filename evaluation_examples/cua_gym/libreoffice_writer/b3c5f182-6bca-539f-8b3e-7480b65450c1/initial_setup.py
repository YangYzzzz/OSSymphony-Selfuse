"""
Initial Setup: Format the merge field 'JoinDate' so it displays as a full date
Task ID: writer_mt_019
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_019'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_mergefield(para, field_name, display_text):
    """
    Insert a MERGEFIELD into a paragraph using Word field codes.
    This creates: { MERGEFIELD JoinDate } with the display_text as cached result.
    No date formatting is applied (raw ISO format).
    """
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Run with fldChar begin
    r_begin = para._element.makeelement(qn('w:r'), {})
    fld_begin = r_begin.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r_begin.append(fld_begin)
    para._element.append(r_begin)

    # Run with instrText
    r_instr = para._element.makeelement(qn('w:r'), {})
    instr = r_instr.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr.text = f' MERGEFIELD {field_name} '
    r_instr.append(instr)
    para._element.append(r_instr)

    # Run with fldChar separate
    r_sep = para._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    r_sep_run = para._element.makeelement(qn('w:r'), {})
    r_sep_run.append(r_sep)
    para._element.append(r_sep_run)

    # Run with display text (cached value)
    r_text = para._element.makeelement(qn('w:r'), {})
    t = r_text.makeelement(qn('w:t'), {qn('xml:space'): 'preserve'})
    t.text = display_text
    r_text.append(t)
    para._element.append(r_text)

    # Run with fldChar end
    r_end = para._element.makeelement(qn('w:r'), {})
    fld_end = r_end.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r_end.append(fld_end)
    para._element.append(r_end)


def create_initial():
    doc = Document()

    # --- Document title ---
    title = doc.add_heading('Welcome to Greenfield Technologies', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sub.add_run('New Employee Welcome Package')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()  # blank line

    # --- Greeting paragraph ---
    p1 = doc.add_paragraph()
    run1 = p1.add_run('Dear ')
    run1.font.size = Pt(11)
    run1.font.name = 'Calibri'

    # Add FirstName merge field
    make_mergefield(p1, 'FirstName', 'Priya')

    run_comma = p1.add_run(',')
    run_comma.font.size = Pt(11)
    run_comma.font.name = 'Calibri'

    # --- Main body paragraphs ---
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(6)
    run2 = p2.add_run(
        'We are thrilled to welcome you to the Greenfield Technologies family! '
        'As a new member of our team, you will find a collaborative and innovative '
        'environment that fosters professional growth and creativity.'
    )
    run2.font.size = Pt(11)
    run2.font.name = 'Calibri'

    # --- The key paragraph with JoinDate merge field ---
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(6)
    run3a = p3.add_run('You officially joined us on ')
    run3a.font.size = Pt(11)
    run3a.font.name = 'Calibri'

    # Insert JoinDate merge field with raw ISO date format (no formatting applied)
    make_mergefield(p3, 'JoinDate', '2025-03-15')

    run3b = p3.add_run('.')
    run3b.font.size = Pt(11)
    run3b.font.name = 'Calibri'

    # --- Department info ---
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_before = Pt(6)
    run4a = p4.add_run('You have been assigned to the ')
    run4a.font.size = Pt(11)
    run4a.font.name = 'Calibri'

    make_mergefield(p4, 'Department', 'Product Engineering')

    run4b = p4.add_run(' department, reporting to ')
    run4b.font.size = Pt(11)
    run4b.font.name = 'Calibri'

    make_mergefield(p4, 'ManagerName', 'David Kim')

    run4c = p4.add_run('.')
    run4c.font.size = Pt(11)
    run4c.font.name = 'Calibri'

    # --- Checklist section ---
    doc.add_heading('Your First Week Checklist', level=2)

    checklist_items = [
        'Complete HR onboarding paperwork (Building A, Room 204)',
        'Set up your workstation and access credentials with IT Support',
        'Attend the new hire orientation session on your first Monday',
        'Review the employee handbook and company policies on the intranet',
        'Meet with your team lead for project briefing and role overview',
        'Enroll in benefits program through the HR portal by end of first week',
    ]
    for item in checklist_items:
        bullet = doc.add_paragraph(item, style='List Bullet')
        for run in bullet.runs:
            run.font.size = Pt(11)
            run.font.name = 'Calibri'

    # --- Contact info ---
    doc.add_paragraph()
    p5 = doc.add_paragraph()
    run5 = p5.add_run(
        'If you have any questions, please reach out to the HR team at '
        'hr@greenfieldtech.com or call extension 4500. We look forward to '
        'working with you!'
    )
    run5.font.size = Pt(11)
    run5.font.name = 'Calibri'

    # --- Closing ---
    doc.add_paragraph()
    p6 = doc.add_paragraph()
    run6 = p6.add_run('Warm regards,')
    run6.font.size = Pt(11)
    run6.font.name = 'Calibri'

    p7 = doc.add_paragraph()
    run7a = p7.add_run('Elena Rodriguez')
    run7a.font.size = Pt(11)
    run7a.font.name = 'Calibri'
    run7a.bold = True

    p8 = doc.add_paragraph()
    run8 = p8.add_run('Director of Human Resources')
    run8.font.size = Pt(11)
    run8.font.name = 'Calibri'

    p9 = doc.add_paragraph()
    run9 = p9.add_run('Greenfield Technologies, Inc.')
    run9.font.size = Pt(11)
    run9.font.name = 'Calibri'
    run9.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
