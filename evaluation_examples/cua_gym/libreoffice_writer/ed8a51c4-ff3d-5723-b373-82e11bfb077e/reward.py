"""
Reward Script: Insert Glossary page before References in engineering_thesis.docx
Task ID: writer_acad_059
Domain: libreoffice_writer
Scoring:
  Component 1: Heading 1 'Glossary' exists before 'References'          (0.30 pts)
  Component 2: Two-column glossary table with 16 rows and correct terms  (0.40 pts)
  Component 3: Header row is bold with light gray background             (0.20 pts)
  Component 4: Page break present before Glossary section               (0.10 pts)
  Total: 1.00
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user/Documents'
TASK_ID = 'writer_acad_059'
FILE_PATH = f'{WORKDIR}/engineering_thesis.docx'

EXPECTED_TERMS = [
    'Annealing', 'Bandwidth', 'Capacitance', 'Dielectric', 'Entropy',
    'Flux', 'Gain', 'Hysteresis', 'Impedance', 'Jitter',
    'Kinetics', 'Latency', 'Modulation', 'Noise Floor', 'Oscillation'
]


def count_page_breaks(doc):
    """Count manual page breaks in the document."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            for br in run.element.findall('.//w:br', ns):
                btype = br.attrib.get(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type'
                )
                if btype == 'page':
                    count += 1
    return count


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document (precondition gate)
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Build list of all Heading 1 paragraphs with indices
    headings = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs)
                if p.style.name == 'Heading 1']
    heading_texts = [t for _, t in headings]
    heading_indices = {t: i for i, t in headings}

    # -----------------------------------------------------------------------
    # Component 1: Heading 1 'Glossary' exists and appears before 'References'
    # -----------------------------------------------------------------------
    try:
        glossary_present = 'Glossary' in heading_texts
        references_present = 'References' in heading_texts

        if glossary_present and references_present:
            g_idx = heading_indices['Glossary']
            r_idx = heading_indices['References']
            if g_idx < r_idx:
                print(f"PASS: Component 1 — Heading 1 'Glossary' found at para {g_idx}, "
                      f"before 'References' at para {r_idx} (0.30 pts)")
                total_score += 0.30
            else:
                print(f"FAIL: Component 1 — 'Glossary' heading found at para {g_idx} but "
                      f"'References' is at para {r_idx}; Glossary must come first")
        elif not glossary_present:
            print("FAIL: Component 1 — No Heading 1 'Glossary' found in document")
        else:
            print("FAIL: Component 1 — 'References' heading not found; cannot verify ordering")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: Glossary table exists with 16 rows, 2 columns, correct terms
    # -----------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 2 — No tables found in document")
        else:
            table = doc.tables[0]
            rows = len(table.rows)
            cols = len(table.columns)

            # Check dimensions: 16 rows (1 header + 15 terms), 2 columns
            if rows != 16:
                print(f"FAIL: Component 2 — Table has {rows} rows, expected 16 (1 header + 15 terms)")
            elif cols != 2:
                print(f"FAIL: Component 2 — Table has {cols} columns, expected 2")
            else:
                # Check header row
                header_term = table.rows[0].cells[0].text.strip()
                header_def = table.rows[0].cells[1].text.strip()
                if header_term.lower() != 'term' or header_def.lower() != 'definition':
                    print(f"FAIL: Component 2 — Header row: expected ['Term','Definition'], "
                          f"found ['{header_term}', '{header_def}']")
                else:
                    # Check all 15 term rows
                    data_terms = [table.rows[i].cells[0].text.strip() for i in range(1, 16)]
                    data_defs = [table.rows[i].cells[1].text.strip() for i in range(1, 16)]

                    missing_terms = [t for t in EXPECTED_TERMS if t not in data_terms]
                    extra_terms = [t for t in data_terms if t not in EXPECTED_TERMS]
                    defs_ok = all(d == '-' for d in data_defs)

                    # Check alphabetical order
                    sorted_terms = sorted(data_terms)
                    is_alphabetical = (data_terms == sorted_terms)

                    if missing_terms:
                        print(f"FAIL: Component 2 — Missing terms: {missing_terms}")
                    elif extra_terms:
                        print(f"FAIL: Component 2 — Unexpected extra terms: {extra_terms}")
                    elif not defs_ok:
                        non_dash = [(i+1, data_defs[i]) for i in range(len(data_defs))
                                    if data_defs[i] != '-']
                        print(f"FAIL: Component 2 — Definition cells not all '-': {non_dash}")
                    elif not is_alphabetical:
                        print(f"FAIL: Component 2 — Terms not in alphabetical order. "
                              f"Actual: {data_terms}")
                    else:
                        print(f"PASS: Component 2 — Table has 16 rows x 2 cols, correct header, "
                              f"all 15 terms alphabetical, definitions are '-' (0.40 pts)")
                        total_score += 0.40
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Header row is bold with light gray background (D3D3D3)
    # -----------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 3 — No tables found; cannot check header formatting")
        else:
            table = doc.tables[0]
            header_row = table.rows[0]

            # Check bold in header cells
            bold_ok = True
            for cell in header_row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip() and not (run.bold or run.font.bold):
                            bold_ok = False

            # Check gray background shading on header cells
            gray_ok = True
            for cell in header_row.cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is None:
                    gray_ok = False
                    break
                shd = tcPr.find(qn('w:shd'))
                if shd is None:
                    gray_ok = False
                    break
                fill = shd.attrib.get(qn('w:fill'), '').upper()
                # Accept D3D3D3 (light gray) or similar gray shades
                if not fill or fill == 'NONE' or fill == 'AUTO':
                    gray_ok = False

            if bold_ok and gray_ok:
                print(f"PASS: Component 3 — Header row is bold and has gray background "
                      f"(fill: {fill}) (0.20 pts)")
                total_score += 0.20
            elif not bold_ok:
                print("FAIL: Component 3 — Header row cells are not bold")
            else:
                print(f"FAIL: Component 3 — Header row background not gray "
                      f"(fill: '{fill}')")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: A page break is present before the Glossary section
    # (initial doc has 10 page breaks; golden has 11 — one added before Glossary)
    # -----------------------------------------------------------------------
    try:
        total_page_breaks = count_page_breaks(doc)
        # The initial doc has 10 page breaks. The task requires a new page before Glossary,
        # so the golden doc should have at least 11 page breaks.
        if total_page_breaks >= 11:
            print(f"PASS: Component 4 — Document has {total_page_breaks} page breaks "
                  f"(>= 11, indicating a new page added before Glossary) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — Document has only {total_page_breaks} page breaks; "
                  f"expected >= 11 (need new page before Glossary)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -----------------------------------------------------------------------
    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
