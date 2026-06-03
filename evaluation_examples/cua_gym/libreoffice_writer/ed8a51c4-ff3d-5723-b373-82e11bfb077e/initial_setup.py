"""
Initial Setup: Engineering thesis document with References section but no Glossary
Task ID: writer_acad_059
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Documents'
TASK_ID = 'engineering_thesis'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading(doc, text, level=1):
    """Add a heading with a specific level."""
    return doc.add_heading(text, level=level)


def add_body_paragraphs(doc, paragraphs):
    """Add multiple body paragraphs."""
    for text in paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(6)


def create_initial():
    # Create Documents directory if not exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set default margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    run = title_para.add_run("Advanced Electromagnetic Systems in Modern Engineering Applications")
    run.font.size = Pt(16)
    run.bold = True

    sub_para = doc.add_paragraph()
    sub_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = sub_para.add_run("A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy in Electrical Engineering")
    run2.font.size = Pt(12)

    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.space_before = Pt(36)
    run3 = author_para.add_run("Dr. Elena Vasquez\nDepartment of Electrical Engineering\nMassachusetts Institute of Technology")
    run3.font.size = Pt(12)

    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.space_before = Pt(24)
    run4 = date_para.add_run("Cambridge, Massachusetts\nSeptember 2024")
    run4.font.size = Pt(11)

    doc.add_page_break()

    # =========================================================================
    # ABSTRACT
    # =========================================================================
    add_heading(doc, "Abstract", level=1)
    add_body_paragraphs(doc, [
        "This thesis presents a comprehensive investigation into the electromagnetic behavior of advanced composite materials under varying thermal and mechanical stress conditions. The research addresses fundamental challenges in designing high-frequency circuit components for aerospace and telecommunications applications.",
        "Through a combination of experimental characterization and computational modeling, we demonstrate that nano-structured ferromagnetic composites exhibit significantly enhanced permeability retention at frequencies exceeding 10 GHz. Our findings reveal a previously unreported correlation between grain boundary density and hysteresis loss mechanisms at elevated temperatures.",
        "The proposed analytical framework integrates classical electromagnetic theory with quantum mechanical corrections to accurately predict material behavior across a wide operational range. Validation against experimental data from 847 test specimens confirms prediction accuracy within 2.3% across all measured parameters.",
        "These results have immediate implications for the design of compact, high-efficiency power converters and low-noise amplification systems. The methodologies developed in this work provide a foundation for next-generation electromagnetic component design.",
    ])

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS (simplified)
    # =========================================================================
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "Abstract ............................................................. ii",
        "List of Figures ..................................................... iv",
        "List of Tables ...................................................... v",
        "1. Introduction ..................................................... 1",
        "2. Literature Review ................................................ 8",
        "3. Theoretical Framework ............................................ 18",
        "4. Experimental Methodology ......................................... 28",
        "5. Results and Analysis ............................................. 41",
        "6. Discussion ....................................................... 58",
        "7. Conclusion ....................................................... 68",
        "References .......................................................... 74",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 1: INTRODUCTION
    # =========================================================================
    add_heading(doc, "1. Introduction", level=1)
    add_heading(doc, "1.1 Background and Motivation", level=2)
    add_body_paragraphs(doc, [
        "The rapid advancement of wireless communication technologies over the past two decades has created an unprecedented demand for electromagnetic components that operate reliably at microwave and millimeter-wave frequencies. Modern 5G networks, radar systems, and satellite communications all require components with exceptional electromagnetic performance characteristics that push the boundaries of conventional material science.",
        "Traditional ferrite-based materials, while adequate for lower frequency applications, exhibit significant performance degradation when subjected to the thermal cycling inherent in high-power applications. The problem is further compounded by the increasing miniaturization trend, which requires components to dissipate more energy per unit volume while maintaining thermal stability.",
        "The work presented in this thesis addresses these challenges through a systematic investigation of nano-structured composite materials. By controlling material composition at the nanoscale, we achieve electromagnetic properties that were previously unattainable with bulk materials, opening new possibilities for compact, high-performance electromagnetic system design.",
    ])

    add_heading(doc, "1.2 Research Objectives", level=2)
    add_body_paragraphs(doc, [
        "The primary objectives of this research are threefold. First, to develop a quantitative understanding of how nano-structured grain boundaries influence electromagnetic loss mechanisms across a wide frequency spectrum. Second, to create predictive models that accurately capture the temperature-dependent behavior of composite materials under operational conditions.",
        "Third, and perhaps most importantly, to translate these fundamental insights into practical design guidelines that engineers can apply directly to component development. This requires bridging the gap between theoretical physics and applied engineering in a way that is both rigorous and practically accessible.",
    ])

    add_heading(doc, "1.3 Thesis Organization", level=2)
    add_body_paragraphs(doc, [
        "This thesis is organized as follows. Chapter 2 provides a comprehensive review of existing literature on electromagnetic composite materials, identifying key gaps in current understanding. Chapter 3 develops the theoretical framework that underpins our experimental and computational work.",
        "Chapter 4 describes the experimental methodology in detail, including material synthesis, characterization techniques, and measurement protocols. Chapter 5 presents the experimental results alongside computational predictions. Chapter 6 discusses the broader implications of our findings and their relationship to existing theory. Chapter 7 summarizes key conclusions and outlines directions for future research.",
    ])

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 2: LITERATURE REVIEW
    # =========================================================================
    add_heading(doc, "2. Literature Review", level=1)
    add_heading(doc, "2.1 Historical Development of Electromagnetic Materials", level=2)
    add_body_paragraphs(doc, [
        "The study of electromagnetic materials dates to the pioneering work of Michael Faraday in the 1830s, whose discoveries laid the groundwork for understanding electromagnetic induction. The subsequent development of Maxwell's equations provided the mathematical framework that still governs our understanding of electromagnetic phenomena today.",
        "The twentieth century saw dramatic advances in material science driven largely by defense and telecommunications applications. The development of radar technology during World War II accelerated research into microwave-absorbing materials, leading to the first generation of practical ferrite components. Subsequent decades brought increasingly sophisticated manufacturing techniques and a deeper understanding of the relationship between material microstructure and electromagnetic properties.",
        "By the 1990s, computational methods had advanced sufficiently to enable reliable modeling of complex electromagnetic structures, transforming component design from an empirical art into a quantitative science. The emergence of nanotechnology in the early 2000s opened entirely new possibilities for material engineering at the atomic scale.",
    ])

    add_heading(doc, "2.2 Ferrite Material Properties", level=2)
    add_body_paragraphs(doc, [
        "Ferrite materials have dominated high-frequency electromagnetic applications for decades due to their favorable combination of high magnetic permeability, low electrical conductivity, and relatively low production cost. The spinel ferrite family, characterized by the chemical formula MFe2O4 where M represents a divalent metal cation, offers particular versatility through compositional modification.",
        "Manganese-zinc (MnZn) ferrites exhibit permeability values typically ranging from 1,000 to 15,000 at frequencies up to approximately 1 MHz. Above this frequency, however, eddy current losses become significant and performance degrades substantially. Nickel-zinc (NiZn) ferrites offer better high-frequency performance, with usable operating ranges extending to several hundred MHz in carefully optimized formulations.",
        "The fundamental limitation of conventional ferrite materials lies in their bulk crystalline structure, where grain boundaries act as defect sites that promote magnetic domain wall pinning and energy dissipation. Reducing grain size to the nanometer scale fundamentally alters the magnetic domain structure, potentially yielding materials with substantially improved high-frequency performance.",
    ])

    add_heading(doc, "2.3 Nano-Structured Magnetic Materials", level=2)
    add_body_paragraphs(doc, [
        "Research into nano-structured magnetic materials began in earnest following the discovery of giant magnetoresistance (GMR) in 1988 by Albert Fert and Peter Grünberg. While GMR itself finds primary application in data storage, the techniques developed for fabricating nanoscale magnetic structures proved broadly applicable to electromagnetic material synthesis.",
        "Studies by Zhang et al. (2008) demonstrated that iron-based nanoparticles with diameters below 20 nm exhibit single-domain behavior, eliminating the domain wall losses that limit bulk material performance. Subsequent work by Hernandez and colleagues (2012) showed that controlled aggregation of such nanoparticles into composite structures could be tuned to achieve specific permeability profiles.",
        "More recent investigations have focused on core-shell nanoparticle architectures, where a magnetic core is surrounded by a non-magnetic shell material. This approach allows independent optimization of magnetic and dielectric properties, enabling composite materials with unprecedented combinations of high permeability and low loss at microwave frequencies.",
    ])

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 3: THEORETICAL FRAMEWORK
    # =========================================================================
    add_heading(doc, "3. Theoretical Framework", level=1)
    add_heading(doc, "3.1 Classical Electromagnetic Theory", level=2)
    add_body_paragraphs(doc, [
        "The electromagnetic behavior of materials is fundamentally governed by Maxwell's equations, which relate electric and magnetic field quantities to charge distributions and current flows. In differential form, these equations provide a complete description of electromagnetic phenomena at macroscopic scales where quantum effects are negligible.",
        "For magnetic materials specifically, the constitutive relationship B = μH, where B is magnetic flux density, μ is permeability, and H is magnetic field intensity, provides the key link between material properties and field behavior. In practice, μ is a complex, frequency-dependent tensor whose components must be determined experimentally for materials with anisotropic structure.",
        "The loss behavior of magnetic materials is characterized by the imaginary component of complex permeability, μ''. The loss tangent tan(δ) = μ''/μ' quantifies the ratio of energy dissipated to energy stored per cycle, providing a convenient figure of merit for comparing material performance in specific applications.",
    ])

    add_heading(doc, "3.2 Quantum Corrections at the Nanoscale", level=2)
    add_body_paragraphs(doc, [
        "When material dimensions approach the nanometer scale, quantum mechanical effects become significant and classical electromagnetic theory requires modification. The exchange interaction between neighboring magnetic moments, which governs the formation of magnetic domains, has a characteristic length scale of approximately 1-10 nm depending on material composition.",
        "Below this exchange length, the formation of domain walls becomes energetically unfavorable, and particles adopt single-domain configurations with uniform magnetization. The switching behavior of single-domain particles is governed by the Stoner-Wohlfarth model, which predicts switching fields as a function of particle geometry and applied field orientation.",
        "For composite materials containing many nanoparticles in various orientations, statistical averaging over the distribution of particle configurations yields effective medium properties that can be predicted using appropriate mixing rules. The Maxwell-Garnett effective medium theory provides a first approximation, though more sophisticated approaches accounting for particle-particle interactions are necessary for accurate predictions at high volume fractions.",
    ])

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 4: EXPERIMENTAL METHODOLOGY
    # =========================================================================
    add_heading(doc, "4. Experimental Methodology", level=1)
    add_heading(doc, "4.1 Material Synthesis", level=2)
    add_body_paragraphs(doc, [
        "Nano-structured composite materials were synthesized using a modified co-precipitation technique that allows precise control of particle size and composition. Starting materials of analytical grade purity (>99.9%) were obtained from Sigma-Aldrich and used without further purification. All synthesis operations were conducted in a nitrogen atmosphere to prevent oxidation of intermediate compounds.",
        "Iron oxide nanoparticles were prepared by dropwise addition of 0.1 M NaOH solution to a mixed aqueous solution of FeCl2 and FeCl3 at room temperature under vigorous mechanical stirring at 1200 rpm. The molar ratio of Fe2+ to Fe3+ was maintained at 1:2 to favor formation of magnetite (Fe3O4) as determined by X-ray diffraction analysis.",
        "Particle size was controlled through reaction temperature and time. Reactions conducted at 25°C for 30 minutes yielded particles with mean diameter of 8.3 ± 1.2 nm, while reactions at 60°C for 60 minutes produced particles of 18.7 ± 2.8 nm as determined by transmission electron microscopy analysis on a minimum of 200 particles per sample.",
    ])

    add_heading(doc, "4.2 Characterization Techniques", level=2)
    add_body_paragraphs(doc, [
        "Structural characterization was performed using X-ray powder diffraction (XRD) on a Bruker D8 Advance diffractometer with Cu Kα radiation (λ = 1.5406 Å). Diffraction patterns were collected over the angular range 20° < 2θ < 80° with a step size of 0.02° and counting time of 2 seconds per step. Phase identification was performed by comparison with ICDD powder diffraction file database entries.",
        "Magnetic properties were measured using a Quantum Design MPMS XL superconducting quantum interference device (SQUID) magnetometer. Magnetization versus applied field measurements were conducted at 300 K and 10 K over field ranges of ±70,000 Oe. Temperature-dependent magnetization was measured under both zero-field-cooled and field-cooled conditions in an applied field of 100 Oe.",
        "High-frequency electromagnetic properties were characterized using a vector network analyzer (Agilent E8363B) operating over the frequency range 0.1-20 GHz. Permittivity and permeability measurements were performed using the transmission/reflection method with samples cut to fit a 7 mm coaxial sample holder. Calibration was performed using the SOLT (short-open-load-through) method.",
    ])

    add_heading(doc, "4.3 Sample Preparation for Electromagnetic Measurements", level=2)
    add_body_paragraphs(doc, [
        "Composite samples for electromagnetic characterization were prepared by dispersing nanoparticles in a paraffin wax matrix at volume fractions ranging from 5% to 45%. Uniform particle dispersion was achieved through ultrasonic mixing at 40 kHz for 30 minutes followed by vacuum degassing to remove entrapped air bubbles.",
        "The composite mixture was cast into toroidal molds with inner diameter 3.04 mm and outer diameter 7.0 mm, conforming to the requirements of the coaxial measurement fixture. After solidification at room temperature over 24 hours, sample faces were machined flat to ensure good contact with the measurement fixture conductors. Sample thickness was maintained at 3.0 ± 0.1 mm to provide sufficient signal-to-noise ratio across the measurement bandwidth.",
    ])

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 5: RESULTS AND ANALYSIS
    # =========================================================================
    add_heading(doc, "5. Results and Analysis", level=1)
    add_heading(doc, "5.1 Structural Characterization Results", level=2)
    add_body_paragraphs(doc, [
        "X-ray diffraction analysis confirmed the formation of phase-pure magnetite (Fe3O4) in all synthesized samples. The characteristic diffraction peaks at 2θ = 30.1°, 35.5°, 43.1°, 53.5°, 57.0°, and 62.6° corresponding to the (220), (311), (400), (422), (511), and (440) planes of the spinel crystal structure were observed in all patterns.",
        "Crystallite size calculated from peak broadening using the Scherrer equation ranged from 7.2 to 22.4 nm depending on synthesis conditions, consistent with the TEM measurements. The lattice parameter refined to a = 8.394 ± 0.003 Å, in good agreement with the accepted value for bulk magnetite (a = 8.396 Å), confirming complete iron oxide formation without significant structural distortion.",
        "Selected area electron diffraction patterns obtained from TEM analysis confirmed the single-crystal nature of individual particles below 12 nm diameter and revealed the development of defect structures in larger particles, suggesting a critical size threshold for the onset of structural complexity.",
    ])

    add_heading(doc, "5.2 Magnetic Properties", level=2)
    add_body_paragraphs(doc, [
        "Room temperature magnetization curves showed characteristic superparamagnetic behavior for particles with diameter below 14 nm, exhibiting no coercivity or remanence within the measurement precision of ±2 Oe. Saturation magnetization (Ms) values ranged from 61 to 78 emu/g, somewhat below the bulk value of 92 emu/g, consistent with the presence of a magnetically disordered surface layer of approximately 1-2 nm thickness.",
        "For samples measured at 10 K, well below the blocking temperature, hysteresis loops were observed with coercive fields (Hc) ranging from 180 to 340 Oe depending on particle size. The ratio of remanent magnetization to saturation magnetization (Mr/Ms) ranged from 0.41 to 0.49, consistent with theoretical predictions for randomly oriented single-domain particles with uniaxial anisotropy.",
        "Blocking temperatures determined from ZFC-FC magnetization curves ranged from 145 K for the smallest (8 nm) particles to 287 K for the largest (22 nm) particles. The observation of blocking temperatures near or above room temperature for larger particles is significant for electromagnetic applications, as it indicates non-superparamagnetic behavior that introduces frequency-independent loss mechanisms.",
    ])

    add_heading(doc, "5.3 High-Frequency Electromagnetic Properties", level=2)
    add_body_paragraphs(doc, [
        "Complex permeability measurements revealed frequency-dependent behavior consistent with magnetic resonance phenomena. For composites containing 8 nm particles at 30% volume fraction, real permeability (μ') remained above 1.8 up to 8 GHz before decreasing sharply near the natural resonance frequency at 11.3 GHz. Imaginary permeability (μ'') peaked at 0.95 at the resonance frequency, yielding a peak loss tangent of 0.52.",
        "Composites prepared with 18 nm particles exhibited notably different frequency behavior. The resonance frequency shifted to 6.7 GHz, consistent with the larger particle size increasing magnetic domain complexity and reducing the exchange-dominated resonance frequency. Peak imaginary permeability reached 1.42, indicating stronger resonance absorption but over a lower frequency range.",
        "Temperature-dependent measurements conducted from -40°C to 150°C revealed remarkable stability in composites containing the smaller (8 nm) particles. Real permeability at 5 GHz varied by less than 8% across this 190°C temperature range, compared to >35% variation for conventional MnZn ferrite reference samples. This enhanced thermal stability is attributed to the superparamagnetic character of the nanoparticles maintaining a thermally driven equilibrium magnetization state.",
    ])

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 6: DISCUSSION
    # =========================================================================
    add_heading(doc, "6. Discussion", level=1)
    add_heading(doc, "6.1 Interpretation of Electromagnetic Results", level=2)
    add_body_paragraphs(doc, [
        "The frequency-dependent electromagnetic properties observed in this study can be understood within the framework of magnetic relaxation theory. The high-frequency permeability behavior of single-domain nanoparticle systems is governed by two primary relaxation mechanisms: Néel relaxation, involving thermal fluctuations of magnetic moment orientation within a fixed particle, and Brownian relaxation, involving physical rotation of the particle within the host medium.",
        "For particles with diameter below approximately 10 nm dispersed in the relatively viscous paraffin wax matrix, Néel relaxation dominates. The characteristic Néel relaxation frequency scales inversely with particle volume, explaining the observed shift in resonance frequency with particle size. The quantitative agreement between predicted and measured resonance frequencies provides strong validation of the single-domain model for these particles.",
        "The anomalously high permeability stability observed across the operational temperature range results from a compensation mechanism unique to superparamagnetic systems. As temperature increases, thermal energy increasingly randomizes magnetic moment orientations, which would normally decrease permeability. However, this is offset by reduced anisotropy energy barriers, which enable more efficient field-driven orientation of magnetic moments. The net result is approximate temperature independence of the complex permeability in a balanced composition regime.",
    ])

    add_heading(doc, "6.2 Comparison with Existing Models", level=2)
    add_body_paragraphs(doc, [
        "Quantitative comparison of experimental results with theoretical predictions was performed using both the Maxwell-Garnett effective medium theory and the more sophisticated Bruggeman symmetric medium approximation. For volume fractions below 20%, both models provided good agreement with measured permeability values, with root mean square errors of 3.1% and 2.7% respectively.",
        "At higher volume fractions (>30%), the Maxwell-Garnett model systematically underestimated permeability by 8-15%, while the Bruggeman model maintained better accuracy (error <5%). This divergence reflects the increased importance of particle-particle interactions at high densities, which are accounted for in the Bruggeman framework through a self-consistent averaging procedure.",
        "The modified effective medium theory incorporating quantum mechanical corrections to the inter-particle exchange interaction, developed in Chapter 3, provided the best overall agreement with a maximum error of 3.2% across all measurement conditions. This validates the theoretical framework and demonstrates the importance of quantum corrections for accurate material modeling at the nanoscale.",
    ])

    add_heading(doc, "6.3 Implications for Device Design", level=2)
    add_body_paragraphs(doc, [
        "The electromagnetic properties demonstrated by the optimized nano-composite materials have several important implications for practical device design. For power electronics applications operating in the 1-5 MHz range, the superior temperature stability offers potential for significant efficiency improvements in switching converters, where core losses are a primary limiting factor for miniaturization.",
        "In wireless communications applications, the tunable resonance frequency of nano-composite materials offers possibilities for frequency-reconfigurable components. By adjusting particle size distribution through synthesis conditions, materials can be tailored for specific frequency bands from 1 GHz to above 15 GHz, covering the major bands used in 5G and beyond applications.",
        "The analysis also reveals important design constraints. The peak loss tangent near the resonance frequency represents a significant limitation for applications requiring low insertion loss across broad bandwidths. Careful material selection to place the resonance frequency well above the operating band, combined with appropriate composite geometry design, is essential for minimizing this loss penalty.",
    ])

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 7: CONCLUSION
    # =========================================================================
    add_heading(doc, "7. Conclusion", level=1)
    add_heading(doc, "7.1 Summary of Findings", level=2)
    add_body_paragraphs(doc, [
        "This thesis has presented a comprehensive investigation of the electromagnetic properties of nano-structured composite materials, with focus on their application in high-frequency circuit components. The key findings can be summarized as follows.",
        "First, nano-structured magnetite composites with controlled particle size in the 8-22 nm range exhibit electromagnetic properties substantially superior to conventional bulk ferrite materials for frequencies above 5 GHz. Specifically, composites containing 8 nm particles maintain permeability above 1.8 up to 8 GHz while exhibiting remarkable thermal stability across the -40°C to 150°C operational range.",
        "Second, the theoretical framework developed in Chapter 3, incorporating quantum mechanical corrections to classical effective medium theory, provides accurate predictions of composite electromagnetic properties across all tested conditions. The maximum prediction error of 3.2% represents a significant improvement over existing models and provides a reliable basis for material optimization in device design workflows.",
        "Third, the experimental methodology developed for simultaneous characterization of structural, magnetic, and electromagnetic properties provides a comprehensive protocol that will support future investigations of this material class.",
    ])

    add_heading(doc, "7.2 Contributions to Knowledge", level=2)
    add_body_paragraphs(doc, [
        "The primary contributions of this work to scientific knowledge are threefold. First, the demonstration of the compensation mechanism responsible for thermal stability in superparamagnetic composites provides a theoretical explanation for previously unexplained experimental observations and suggests a general design principle for thermally stable electromagnetic materials.",
        "Second, the modified effective medium theory incorporating nanoscale quantum corrections advances the state of the art in electromagnetic material modeling. Third, the comprehensive experimental database of structure-property relationships for magnetite nano-composites provides a valuable reference for the materials science and electromagnetic communities.",
    ])

    add_heading(doc, "7.3 Future Work", level=2)
    add_body_paragraphs(doc, [
        "Several promising directions for future research emerge from this work. The extension of the experimental characterization to frequencies above 20 GHz, enabled by waveguide measurement techniques, would provide insight into material behavior in the millimeter-wave bands increasingly important for 5G and satellite communications.",
        "Investigation of alternative matrix materials, including polymer composites with tailored dielectric properties, could enable simultaneous optimization of both permittivity and permeability profiles for specific applications. Additionally, the scalability of the synthesis technique to industrial production volumes remains to be demonstrated, representing an important bridge between laboratory results and practical applications.",
        "Finally, the integration of nano-composite materials with conventional printed circuit board manufacturing processes would enable direct incorporation of electromagnetic functionality into circuit substrates, potentially transforming the design of compact RF and microwave systems.",
    ])

    doc.add_page_break()

    # =========================================================================
    # REFERENCES (no Glossary before this - that's the task to add)
    # =========================================================================
    add_heading(doc, "References", level=1)
    references = [
        "[1] A. Fert and P. Grünberg, \"Giant magnetoresistance of Fe/Cr magnetic superlattices,\" Physical Review Letters, vol. 61, no. 21, pp. 2472-2475, 1988.",
        "[2] Z. Zhang, Y. Liu, and T. Wang, \"Superparamagnetic iron oxide nanoparticles: synthesis, characterization, and biomedical applications,\" Advanced Materials, vol. 20, no. 6, pp. 1030-1050, 2008.",
        "[3] R. Hernandez, A. Salas, and M. Garcia, \"Core-shell magnetic nanoparticles with tunable electromagnetic properties for microwave applications,\" Journal of Applied Physics, vol. 112, no. 7, pp. 073912, 2012.",
        "[4] J. Maxwell, A Treatise on Electricity and Magnetism. Oxford: Clarendon Press, 1873.",
        "[5] L. Néel, \"Théorie du traînage magnétique des ferromagnétiques en grains fins avec applications aux terres cuites,\" Annales de Géophysique, vol. 5, pp. 99-136, 1949.",
        "[6] E. Stoner and E. Wohlfarth, \"A mechanism of magnetic hysteresis in heterogeneous alloys,\" Philosophical Transactions of the Royal Society A, vol. 240, pp. 599-642, 1948.",
        "[7] R. Bruggeman, \"Berechnung verschiedener physikalischer Konstanten von heterogenen Substanzen,\" Annalen der Physik, vol. 24, pp. 636-679, 1935.",
        "[8] J. Garnett, \"Colours in metal glasses and in metallic films,\" Philosophical Transactions of the Royal Society A, vol. 203, pp. 385-420, 1904.",
        "[9] W. Brown, \"Thermal fluctuations of a single-domain particle,\" Physical Review, vol. 130, no. 5, pp. 1677-1686, 1963.",
        "[10] S. Chikazumi, Physics of Ferromagnetism. New York: Oxford University Press, 1997.",
        "[11] A. Goldman, Modern Ferrite Technology, 2nd ed. New York: Springer, 2006.",
        "[12] D. Jiles, Introduction to Magnetism and Magnetic Materials. London: Chapman and Hall, 1991.",
        "[13] C. Kittel, Introduction to Solid State Physics, 8th ed. New York: Wiley, 2005.",
        "[14] R. O'Handley, Modern Magnetic Materials: Principles and Applications. New York: Wiley-Interscience, 2000.",
        "[15] T. Spaldin, Magnetic Materials: Fundamentals and Device Applications. Cambridge: Cambridge University Press, 2003.",
        "[16] H. Ibach and H. Lüth, Solid-State Physics: An Introduction to Principles of Materials Science. Berlin: Springer, 2009.",
        "[17] P. Lorrain and D. Corson, Electromagnetic Fields and Waves, 2nd ed. San Francisco: W.H. Freeman, 1970.",
        "[18] D. Pozar, Microwave Engineering, 4th ed. New York: Wiley, 2012.",
        "[19] C. Balanis, Advanced Engineering Electromagnetics, 2nd ed. New York: Wiley, 2012.",
        "[20] K. Gupta, R. Garg, I. Bahl, and P. Trivedi, Microstrip Lines and Slotlines, 3rd ed. Norwood: Artech House, 2013.",
        "[21] B. Lax and K. Button, Microwave Ferrites and Ferrimagnetics. New York: McGraw-Hill, 1962.",
        "[22] A. Tebble and D. Craik, Magnetic Domains. New York: Wiley-Interscience, 1969.",
        "[23] W. Heisenberg, \"Zur Theorie des Ferromagnetismus,\" Zeitschrift für Physik, vol. 49, pp. 619-636, 1928.",
        "[24] M. Knobel, W. Nunes, L. Socolovsky, E. De Biasi, J. Vargas, and J. Denardin, \"Superparamagnetism and other magnetic features in granular materials,\" Journal of Nanoscience and Nanotechnology, vol. 8, pp. 2836-2857, 2008.",
        "[25] Q. Pankhurst, J. Connolly, S. Jones, and J. Dobson, \"Applications of magnetic nanoparticles in biomedicine,\" Journal of Physics D: Applied Physics, vol. 36, pp. R167-R181, 2003.",
    ]
    for ref in references:
        p = doc.add_paragraph(ref, style='Normal')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.5)
        # Add hanging indent
        p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
