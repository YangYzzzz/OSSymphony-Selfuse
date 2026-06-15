"""
Reward Script: Add three bibliography entries and in-text cross-references
Task ID: osworld_writer_bibliography_crossref_007
Domain: libreoffice_writer
Scoring:
  Component 1: In-text citation (Chen, 2020) present in paragraph 3 of body (0.20 pts)
  Component 2: In-text citation (Davis, 2021) present in paragraph 4 of body (0.20 pts)
  Component 3: In-text citation (Evans, 2022) present in paragraph 5 of body (0.20 pts)
  Component 4: Bibliography entry for Chen, W. (2020) added (0.13 pts)
  Component 5: Bibliography entry for Davis, R. (2021) added (0.14 pts)
  Component 6: Bibliography entry for Evans, T. (2022) added (0.13 pts)
  Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_bibliography_crossref_007'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Add three bibliography entries and in-text cross-references.
    - (Chen, 2020) -> paragraph 2 (3rd paragraph = index 3 in doc.paragraphs, 2nd body paragraph after abstract)
    - (Davis, 2021) -> paragraph 3 (4th paragraph = index 4, 3rd body paragraph)
    - (Evans, 2022) -> paragraph 4 (5th paragraph = index 5, 4th body paragraph)
    Bibliography entries to be appended after existing 2 entries.
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather paragraph texts upfront
    paras = doc.paragraphs
    num_paras = len(paras)
    print(f"INFO: Total paragraphs in document: {num_paras}")

    # Find body paragraphs (Normal style paragraphs after the Abstract heading)
    # Para 0: Title
    # Para 1: Abstract Heading
    # Para 2: 1st body paragraph (AI overview)
    # Para 3: 2nd body paragraph (Machine Learning) -> should have (Chen, 2020)
    # Para 4: 3rd body paragraph (Neural Networks)  -> should have (Davis, 2021)
    # Para 5: 4th body paragraph (Deep Learning)    -> should have (Evans, 2022)
    # Para 6: 5th body paragraph (Society)
    # Para 7: Bibliography heading
    # Para 8+: bibliography entries

    # Component 1: In-text citation (Chen, 2020) in paragraph 2 of body text
    # This is the 2nd body paragraph (index 3 in doc.paragraphs = 'Machine learning...')
    try:
        # The task says "paragraph 2" meaning the 2nd body paragraph
        # doc.paragraphs[3] is the Machine Learning paragraph
        p3_text = paras[3].text if num_paras > 3 else ""
        if "(Chen, 2020)" in p3_text:
            print(f"PASS: Component 1 — '(Chen, 2020)' found in paragraph 3 (Machine Learning) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — '(Chen, 2020)' NOT found in paragraph 3. Text: {p3_text[-100:]!r}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: In-text citation (Davis, 2021) in paragraph 3 of body text
    # This is the 3rd body paragraph (index 4 in doc.paragraphs = 'Neural network...')
    try:
        p4_text = paras[4].text if num_paras > 4 else ""
        if "(Davis, 2021)" in p4_text:
            print(f"PASS: Component 2 — '(Davis, 2021)' found in paragraph 4 (Neural Networks) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — '(Davis, 2021)' NOT found in paragraph 4. Text: {p4_text[-100:]!r}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: In-text citation (Evans, 2022) in paragraph 4 of body text
    # This is the 4th body paragraph (index 5 in doc.paragraphs = 'Deep learning...')
    try:
        p5_text = paras[5].text if num_paras > 5 else ""
        if "(Evans, 2022)" in p5_text:
            print(f"PASS: Component 3 — '(Evans, 2022)' found in paragraph 5 (Deep Learning) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — '(Evans, 2022)' NOT found in paragraph 5. Text: {p5_text[-100:]!r}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Bibliography entry for Chen, W. (2020) added
    # Search all paragraphs for this entry using any() to avoid loop-assigned booleans
    try:
        chen_entry_count = sum(
            1 for p in paras
            if "Chen" in p.text and "2020" in p.text and ("Machine Learning" in p.text or "Tech Press" in p.text)
        )
        if chen_entry_count > 0:
            print(f"PASS: Component 4 — Bibliography entry 'Chen, W. (2020). Machine Learning Applications. Tech Press.' found (0.13 pts)")
            total_score += 0.13
        else:
            print(f"FAIL: Component 4 — Bibliography entry for Chen (2020) NOT found in document")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Bibliography entry for Davis, R. (2021) added
    try:
        davis_entry_count = sum(
            1 for p in paras
            if "Davis" in p.text and "2021" in p.text and ("Neural Networks" in p.text or "AI Books" in p.text)
        )
        if davis_entry_count > 0:
            print(f"PASS: Component 5 — Bibliography entry 'Davis, R. (2021). Neural Networks Guide. AI Books.' found (0.14 pts)")
            total_score += 0.14
        else:
            print(f"FAIL: Component 5 — Bibliography entry for Davis (2021) NOT found in document")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Bibliography entry for Evans, T. (2022) added
    try:
        evans_entry_count = sum(
            1 for p in paras
            if "Evans" in p.text and "2022" in p.text and ("Deep Learning" in p.text or "Science Press" in p.text)
        )
        if evans_entry_count > 0:
            print(f"PASS: Component 6 — Bibliography entry 'Evans, T. (2022). Deep Learning Fundamentals. Science Press.' found (0.13 pts)")
            total_score += 0.13
        else:
            print(f"FAIL: Component 6 — Bibliography entry for Evans (2022) NOT found in document")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path on the VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
