"""
Initial Setup: Technical report with existing bibliography (2 entries), no in-text citations for new entries.
Task ID: osworld_writer_bibliography_crossref_007
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_bibliography_crossref_007'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Advances in Artificial Intelligence: A Survey', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Abstract heading
    doc.add_heading('Abstract', level=1)

    # Paragraph 1 - Introduction
    doc.add_paragraph(
        'Artificial intelligence (AI) has undergone remarkable transformation over the past two decades, '
        'evolving from rule-based systems to sophisticated learning architectures capable of surpassing '
        'human performance on numerous benchmarks (Anderson, 2018). The proliferation of large-scale '
        'datasets and computational advances has catalyzed progress across diverse domains including '
        'natural language processing, computer vision, and autonomous systems. This report surveys '
        'key developments in the field, highlighting methodological innovations and their practical '
        'applications in modern computing environments.'
    )

    # Paragraph 2 - Machine Learning
    doc.add_paragraph(
        'Machine learning has emerged as the dominant paradigm in AI research, enabling systems to '
        'learn patterns from data without explicit programming. Supervised learning techniques have '
        'achieved state-of-the-art results in classification and regression tasks across multiple '
        'industry sectors. The development of gradient-based optimization methods has made it feasible '
        'to train models with billions of parameters efficiently. Recent advances in transfer learning '
        'have further reduced the data requirements for specialized applications, democratizing access '
        'to high-performance AI across organizations of varying sizes.'
    )

    # Paragraph 3 - Neural Networks
    doc.add_paragraph(
        'Neural network architectures have become increasingly sophisticated, with attention mechanisms '
        'and transformer-based models revolutionizing sequence modeling tasks. Recurrent neural networks '
        'and long short-term memory units established foundational capabilities for processing sequential '
        'data, while convolutional architectures demonstrated exceptional performance in spatial pattern '
        'recognition. The integration of residual connections and normalization techniques has enabled '
        'training of very deep networks without gradient degradation. Ensemble methods combining multiple '
        'network architectures have proven particularly effective for robust prediction in high-stakes '
        'applications (Brown, 2019).'
    )

    # Paragraph 4 - Deep Learning
    doc.add_paragraph(
        'Deep learning frameworks have accelerated research productivity by providing high-level '
        'abstractions for model construction and automatic differentiation for gradient computation. '
        'The availability of GPU-accelerated computing infrastructure has reduced training times from '
        'weeks to hours for complex architectures. Reinforcement learning combined with deep neural '
        'networks has demonstrated superhuman performance in game-playing environments, suggesting '
        'promising pathways for autonomous decision-making systems. Generative models including '
        'variational autoencoders and generative adversarial networks have opened new avenues for '
        'data synthesis and creative applications.'
    )

    # Paragraph 5 - Conclusion
    doc.add_paragraph(
        'The rapid advancement of AI capabilities presents both significant opportunities and important '
        'challenges for society. Ethical considerations surrounding algorithmic bias, data privacy, and '
        'transparency must be addressed alongside technical progress. Interdisciplinary collaboration '
        'between computer scientists, domain experts, ethicists, and policymakers will be essential for '
        'responsible AI deployment. Future research directions include energy-efficient architectures, '
        'continual learning systems, and robust methods for uncertainty quantification. The field '
        'continues to evolve at an extraordinary pace, promising transformative impacts across virtually '
        'all sectors of human endeavor.'
    )

    # Bibliography section
    doc.add_heading('Bibliography', level=1)

    # Existing entry 1
    bib1 = doc.add_paragraph()
    run1 = bib1.add_run('Anderson, J. (2018). ')
    run1.bold = False
    run2 = bib1.add_run('Foundations of Artificial Intelligence.')
    run2.italic = True
    run3 = bib1.add_run(' Academic Press.')

    # Existing entry 2
    bib2 = doc.add_paragraph()
    run4 = bib2.add_run('Brown, K. (2019). ')
    run5 = bib2.add_run('Ensemble Methods in Statistical Learning.')
    run5.italic = True
    run6 = bib2.add_run(' University Press.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
