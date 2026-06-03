"""
Initial Setup: Brand Guidelines Document with Section Headings
Task ID: writer_mktg_040
Domain: libreoffice_writer

Creates a 20-page brand guidelines document with:
- Cover page
- Plain text navigation list on page 2 (NOT linked)
- Major section headings at pages 3, 6, 9, 13, 17
- No bookmarks or hyperlinks (agent must add them)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_040'
DESKTOP = '/home/user/Desktop'
OUTPUT = f'{DESKTOP}/brand_guidelines_final.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add a manual page break paragraph."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)
    return para


def add_filler_content(doc, num_paragraphs, section_theme):
    """Add realistic filler paragraphs for a section."""
    content_pool = {
        'logo': [
            "The logo is the cornerstone of our brand identity. It must be reproduced faithfully and consistently across all media.",
            "Our primary logo consists of the wordmark and the icon, which should always appear together unless specifically stated otherwise.",
            "Minimum size requirements: Digital — 32px height; Print — 0.5 inches height. Never reproduce the logo smaller than these minimums.",
            "The logo must always be surrounded by a clear space equal to the height of the capital letter 'B' in the wordmark on all sides.",
            "Approved color variations: Full-color on white, Full-color on dark backgrounds, One-color black, One-color white (reversed).",
            "Do not stretch, compress, rotate, or otherwise distort the logo. Do not add drop shadows, outlines, or other effects.",
            "When placing the logo on photographic backgrounds, ensure sufficient contrast. Use the reversed white version on dark images.",
            "The logo should never be placed on busy or highly textured backgrounds that reduce legibility.",
            "For co-branding scenarios, maintain equal visual weight between our logo and partner logos. Separate with a vertical rule.",
            "Digital formats accepted: SVG (preferred for web), PNG with transparency, EPS (for print production).",
            "Request official logo files from the brand team. Never recreate the logo from scratch or approximate it.",
            "Regional logo variations must be approved by the global brand team before use in any external materials.",
        ],
        'colors': [
            "Our color palette has been carefully selected to reflect our brand values of trust, innovation, and accessibility.",
            "Primary Blue — Hex #1E3A8A, RGB 30/58/138, CMYK 78/58/0/46, Pantone 2756 C. Use for primary actions, headlines, key UI elements.",
            "Secondary Teal — Hex #0D9488, RGB 13/148/136, CMYK 91/0/8/42, Pantone 326 C. Use for accents, data visualization, highlights.",
            "Neutral Gray — Hex #6B7280, RGB 107/114/128, CMYK 16/11/0/50. Use for body text, secondary content, borders.",
            "Background White — Hex #FFFFFF and Light Gray #F9FAFB. Use for page backgrounds and card surfaces.",
            "Alert Red — Hex #DC2626, RGB 220/38/38. Use sparingly for error states and critical warnings only.",
            "Color accessibility: All text/background combinations must meet WCAG 2.1 AA contrast ratio of 4.5:1 minimum.",
            "Do not use colors outside the approved palette without brand team approval. Custom tints require written sign-off.",
            "In print production, always specify Pantone colors to ensure color consistency across vendors and substrates.",
            "Digital applications should use hex values as specified. Never rely on screen matching for color-critical print work.",
            "Gradient usage: Approved gradients transition from Primary Blue to Secondary Teal, top-to-bottom or left-to-right only.",
        ],
        'typography': [
            "Typography is a fundamental pillar of our visual identity. Consistent type application reinforces our professional tone.",
            "Primary Typeface: Inter. Use for all digital interfaces, presentations, and most print applications. Free via Google Fonts.",
            "Secondary Typeface: Georgia. Use for long-form editorial content, case studies, and thought leadership articles.",
            "Accent Typeface: Space Mono. Use sparingly for code snippets, technical documentation, and data labels.",
            "Heading scale: H1 — 48px/3rem, H2 — 36px/2.25rem, H3 — 30px/1.875rem, H4 — 24px/1.5rem, H5 — 20px/1.25rem.",
            "Body text: 16px/1rem with 1.6 line-height for optimal readability on screens. Print body text: 10-11pt.",
            "Letter-spacing: Headlines — -0.025em, Body — 0em, All-caps labels — 0.05em minimum.",
            "Never set body text in all capitals. Reserve all-caps for short labels, navigation items, and button text only.",
            "Minimum font sizes: Body text 10pt print / 14px digital. Captions/footnotes 8pt print / 12px digital.",
            "Font pairing rules: Inter headlines + Georgia body for editorial. Inter throughout for business documents.",
            "Web font loading: Implement font-display: swap for performance. Always specify system font fallbacks.",
            "Licensed fonts must not be shared externally. Freelancers and agencies must purchase their own licenses.",
        ],
        'voice': [
            "Our brand voice is the consistent personality we project in all written communications, both internal and external.",
            "Brand personality pillars: Confident without arrogance, Expert without jargon, Human without being casual, Innovative without being obscure.",
            "Tone adaptation: Our tone shifts based on context while the voice remains constant. Formal for legal documents, warmer for community posts.",
            "Writing for clarity: Use active voice. Prefer short sentences (15-20 words max). Avoid passive constructions.",
            "Inclusive language: Use gender-neutral terms. Avoid idioms, metaphors, and colloquialisms that don't translate well globally.",
            "Technical writing: Define acronyms on first use. Write numbers one through nine as words; use numerals for 10 and above.",
            "Punctuation: Oxford comma required. Em dashes for parenthetical clauses — with spaces. No exclamation marks in formal copy.",
            "Avoid: Buzzwords (synergy, leverage, pivot, disrupt), redundant phrases, corporate jargon, and filler words.",
            "Headlines should be action-oriented and benefit-focused. Lead with the most important information (inverted pyramid).",
            "Error messages: Be specific about what went wrong and how to fix it. Never blame the user. Keep a helpful tone.",
            "CTAs (calls to action): Use specific action verbs. 'Download the report' not 'Click here'. Max 5 words.",
            "Localization: All content for non-English markets must be professionally translated. Machine translation is not approved.",
        ],
        'imagery': [
            "Photography is one of the most powerful tools for expressing our brand character and connecting emotionally with audiences.",
            "Photography style: Authentic, natural light. Real people in real situations. Avoid overly posed or staged compositions.",
            "Subject matter: Diverse representation is essential. Feature people of varied ages, ethnicities, body types, and abilities.",
            "Composition: Rule of thirds preferred. Leave breathing room. Avoid centered, symmetrical compositions for dynamic content.",
            "Color grading: Subtle, warm-neutral tone. Avoid heavy filters, extreme contrasts, or trendy Instagram-style treatments.",
            "Prohibited: Stock photo clichés (handshakes, thumbs-up, cheesy smiles), low-resolution images under 150 DPI for print.",
            "Image sizing for print: Full-bleed — minimum 300 DPI at final size. Half-page — 300 DPI. Quarter-page — 200 DPI minimum.",
            "Web image optimization: JPEG for photos (quality 80-85), WebP with JPEG fallback for modern browsers.",
            "Alt text requirement: All images in digital properties must have descriptive alt text for screen readers.",
            "Illustration style: Flat design with our primary color palette. Line weight 2px. Icons from our approved icon library only.",
            "Video thumbnails: Must follow still photography guidelines. Include brand overlay in lower-right corner.",
            "Image licensing: All photography must be licensed for commercial use. Maintain records of licenses in the asset management system.",
        ],
    }

    pool = content_pool.get(section_theme, content_pool['logo'])
    for i in range(num_paragraphs):
        text = pool[i % len(pool)]
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(6)
    return


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set page size to standard letter
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # =========================================================
    # PAGE 1: COVER PAGE
    # =========================================================
    # Brand name
    cover_title = doc.add_paragraph()
    cover_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cover_title.paragraph_format.space_before = Pt(72)
    run = cover_title.add_run("NOVA DYNAMICS")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Brand Guidelines")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    version_para = doc.add_paragraph()
    version_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    version_para.paragraph_format.space_before = Pt(12)
    run = version_para.add_run("Version 3.2 — Final")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run("January 2025")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()
    doc.add_paragraph()

    confidential = doc.add_paragraph()
    confidential.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = confidential.add_run("CONFIDENTIAL — Internal Use Only")
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    dept_para = doc.add_paragraph()
    dept_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = dept_para.add_run("Brand & Communications Department")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Page break → page 2
    add_page_break(doc)

    # =========================================================
    # PAGE 2: PLAIN TEXT NAVIGATION LIST (not linked)
    # =========================================================
    nav_heading = doc.add_paragraph("Table of Contents")
    nav_heading.style = doc.styles['Heading 2']
    nav_heading.paragraph_format.space_after = Pt(12)

    nav_intro = doc.add_paragraph(
        "This document provides comprehensive brand guidelines for all Nova Dynamics communications. "
        "Use the section list below to navigate to each topic."
    )
    nav_intro.paragraph_format.space_after = Pt(18)

    # Section names as plain text list (NOT hyperlinked — agent must do this)
    section_names = [
        "Logo Usage",
        "Color Palette",
        "Typography",
        "Brand Voice & Tone",
        "Photography & Imagery",
    ]
    for name in section_names:
        item = doc.add_paragraph()
        run = item.add_run(name)
        run.font.size = Pt(12)
        item.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()
    doc.add_paragraph(
        "For questions about brand usage, contact the Brand & Communications team at brand@novadynamics.com."
    ).paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "All brand assets are available in the Nova Dynamics Asset Library at assets.novadynamics.com."
    )

    # Page break → page 3
    add_page_break(doc)

    # =========================================================
    # PAGES 3–5: LOGO USAGE SECTION
    # =========================================================
    logo_heading = doc.add_paragraph("Logo Usage")
    logo_heading.style = doc.styles['Heading 1']
    logo_heading.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        "The Nova Dynamics logo represents our commitment to precision engineering and forward-thinking solutions. "
        "This section establishes the rules for proper logo application across all platforms and materials."
    )
    add_filler_content(doc, 10, 'logo')

    # Page break → page 4
    add_page_break(doc)

    add_filler_content(doc, 6, 'logo')

    doc.add_paragraph("Logo Misuse Examples").style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Heading 2']
    doc.add_paragraph(
        "The following examples illustrate common logo misuses that must be avoided at all times. "
        "When in doubt, consult the brand team before proceeding with any logo placement."
    )
    add_filler_content(doc, 3, 'logo')

    # Page break → page 5
    add_page_break(doc)

    doc.add_paragraph("Logo on Backgrounds").style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Heading 2']
    add_filler_content(doc, 5, 'logo')

    # Page break → page 6
    add_page_break(doc)

    # =========================================================
    # PAGES 6–8: COLOR PALETTE SECTION
    # =========================================================
    color_heading = doc.add_paragraph("Color Palette")
    color_heading.style = doc.styles['Heading 1']
    color_heading.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        "Our color system has been engineered to convey trust, innovation, and clarity. "
        "Every color in our palette carries strategic intent and must be applied consistently."
    )
    add_filler_content(doc, 9, 'colors')

    # Page break → page 7
    add_page_break(doc)

    add_filler_content(doc, 5, 'colors')

    doc.add_paragraph("Color in Digital Applications").style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Heading 2']
    add_filler_content(doc, 4, 'colors')

    # Page break → page 8
    add_page_break(doc)

    doc.add_paragraph("Color Accessibility").style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Heading 2']
    add_filler_content(doc, 6, 'colors')

    # Page break → page 9
    add_page_break(doc)

    # =========================================================
    # PAGES 9–12: TYPOGRAPHY SECTION
    # =========================================================
    typo_heading = doc.add_paragraph("Typography")
    typo_heading.style = doc.styles['Heading 1']
    typo_heading.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        "Our typographic system balances professionalism with approachability. "
        "Consistent type application is essential for a unified brand experience across touchpoints."
    )
    add_filler_content(doc, 10, 'typography')

    # Page break → page 10
    add_page_break(doc)

    add_filler_content(doc, 6, 'typography')

    doc.add_paragraph("Type in Practice").style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Heading 2']
    add_filler_content(doc, 4, 'typography')

    # Page break → page 11
    add_page_break(doc)

    doc.add_paragraph("Type Hierarchy Examples").style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Heading 2']
    add_filler_content(doc, 6, 'typography')

    # Page break → page 12
    add_page_break(doc)

    add_filler_content(doc, 6, 'typography')

    # Page break → page 13
    add_page_break(doc)

    # =========================================================
    # PAGES 13–16: BRAND VOICE & TONE SECTION
    # =========================================================
    voice_heading = doc.add_paragraph("Brand Voice & Tone")
    voice_heading.style = doc.styles['Heading 1']
    voice_heading.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        "How we speak is as important as what we say. Our voice is the human expression of our brand, "
        "and our tone adapts to context while remaining true to our core character."
    )
    add_filler_content(doc, 10, 'voice')

    # Page break → page 14
    add_page_break(doc)

    add_filler_content(doc, 6, 'voice')

    doc.add_paragraph("Writing for Different Channels").style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Heading 2']
    add_filler_content(doc, 4, 'voice')

    # Page break → page 15
    add_page_break(doc)

    doc.add_paragraph("Tone of Voice Examples").style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Heading 2']
    add_filler_content(doc, 6, 'voice')

    # Page break → page 16
    add_page_break(doc)

    add_filler_content(doc, 5, 'voice')

    # Page break → page 17
    add_page_break(doc)

    # =========================================================
    # PAGES 17–20: PHOTOGRAPHY & IMAGERY SECTION
    # =========================================================
    imagery_heading = doc.add_paragraph("Photography & Imagery")
    imagery_heading.style = doc.styles['Heading 1']
    imagery_heading.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        "Visual storytelling through photography and illustration is central to how we communicate our brand values. "
        "Our imagery must feel authentic, diverse, and purposeful."
    )
    add_filler_content(doc, 10, 'imagery')

    # Page break → page 18
    add_page_break(doc)

    add_filler_content(doc, 6, 'imagery')

    doc.add_paragraph("Illustration Guidelines").style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Heading 2']
    add_filler_content(doc, 4, 'imagery')

    # Page break → page 19
    add_page_break(doc)

    doc.add_paragraph("Asset Management").style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Heading 2']
    add_filler_content(doc, 6, 'imagery')

    # Page break → page 20
    add_page_break(doc)

    # Final page: Contact and revision history
    doc.add_paragraph("Revision History & Contacts").style = doc.styles['Heading 2']
    doc.add_paragraph(
        "This document is maintained by the Brand & Communications Department. "
        "For questions, corrections, or usage approvals, contact the brand team."
    )

    contacts_data = [
        ("Brand Manager", "Jennifer Zhao", "j.zhao@novadynamics.com"),
        ("Creative Director", "Marcus Sullivan", "m.sullivan@novadynamics.com"),
        ("Digital Brand Lead", "Priya Nair", "p.nair@novadynamics.com"),
        ("Print Production", "David Kowalski", "d.kowalski@novadynamics.com"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Role'
    hdr[1].text = 'Name'
    hdr[2].text = 'Email'
    for role, name, email in contacts_data:
        row = table.add_row().cells
        row[0].text = role
        row[1].text = name
        row[2].text = email

    doc.add_paragraph()
    revision_table_heading = doc.add_paragraph("Revision History")
    revision_table_heading.style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Heading 2']

    revisions = [
        ("1.0", "March 2022", "Initial release"),
        ("2.0", "November 2022", "Added digital color specs and accessibility guidelines"),
        ("2.5", "April 2023", "Updated photography direction; new imagery examples"),
        ("3.0", "September 2023", "Typography overhaul; Inter typeface adoption"),
        ("3.1", "June 2024", "Voice & Tone section expanded; new examples added"),
        ("3.2", "January 2025", "Final review; photography guidelines updated"),
    ]
    rev_table = doc.add_table(rows=1, cols=3)
    rev_table.style = 'Table Grid'
    rev_hdr = rev_table.rows[0].cells
    rev_hdr[0].text = 'Version'
    rev_hdr[1].text = 'Date'
    rev_hdr[2].text = 'Changes'
    for ver, date, changes in revisions:
        row = rev_table.add_row().cells
        row[0].text = ver
        row[1].text = date
        row[2].text = changes

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
