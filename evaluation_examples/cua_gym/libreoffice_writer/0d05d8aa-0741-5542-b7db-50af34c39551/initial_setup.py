"""
Initial Setup: Safety manual with a WARNING paragraph that needs callout box formatting.
Task ID: writer_rd_047
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_047'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_run_font(run, name="Liberation Serif", size=11, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold


def add_heading_paragraph(doc, text, level=1):
    """Add a heading styled as a simple bold larger paragraph for consistency."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if level == 0:
        run.font.size = Pt(18)
        run.bold = True
    elif level == 1:
        run.font.size = Pt(14)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(12)
        run.bold = True
    run.font.name = "Liberation Serif"
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_body_paragraph(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, "Liberation Serif", 11)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    return para


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = "Liberation Serif"
    font.size = Pt(11)

    # Page margins
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # === PAGE 1 ===
    # Title
    add_heading_paragraph(doc, "Meridian Industries Workplace Safety Manual", level=0)
    add_heading_paragraph(doc, "Revision 4.2 - March 2025", level=2)

    add_body_paragraph(doc, (
        "This manual establishes the safety policies, procedures, and guidelines for all employees "
        "of Meridian Industries. Compliance with these standards is mandatory for all personnel "
        "working in manufacturing, laboratory, warehouse, and office environments. Failure to "
        "follow established safety protocols may result in disciplinary action."
    ))

    add_heading_paragraph(doc, "1. General Safety Principles", level=1)

    add_body_paragraph(doc, (
        "All employees must complete the mandatory safety orientation within their first week "
        "of employment. This orientation covers fire evacuation procedures, emergency contact "
        "numbers, first aid station locations, and the proper use of personal protective equipment "
        "(PPE). Refresher training is required annually for all staff members."
    ))

    add_body_paragraph(doc, (
        "Each department maintains a dedicated safety officer responsible for conducting monthly "
        "inspections, documenting hazards, and coordinating with the Environmental Health and "
        "Safety (EHS) division. Safety officers are identified by their green badge lanyards "
        "and are authorized to halt operations if an imminent danger is identified."
    ))

    add_body_paragraph(doc, (
        "Incident reporting is a critical component of our safety culture. Any workplace injury, "
        "near-miss event, or equipment malfunction must be reported within 24 hours using Form "
        "SR-100, available on the company intranet. Anonymous reporting is also available through "
        "the EHS hotline at extension 4500."
    ))

    add_heading_paragraph(doc, "2. Hazard Communication", level=1)

    add_body_paragraph(doc, (
        "Meridian Industries follows the Globally Harmonized System (GHS) for chemical labeling "
        "and Safety Data Sheets (SDS). All containers of hazardous materials must display proper "
        "GHS labels including pictograms, signal words, and hazard statements. Secondary containers "
        "must also be labeled before leaving the dispensing area."
    ))

    add_body_paragraph(doc, (
        "Safety Data Sheets are maintained in both digital and physical formats. Digital copies "
        "are accessible through the EHS portal under the Chemical Inventory section. Physical "
        "binders are located at each work station where chemicals are stored or used. Employees "
        "must review the SDS before handling any unfamiliar substance."
    ))

    add_body_paragraph(doc, (
        "Chemical spill response kits are positioned at every entrance to laboratories and "
        "chemical storage areas. Each kit contains absorbent pads, neutralizing agents, protective "
        "gloves, goggles, and disposal bags. Spill response procedures are posted on laminated "
        "cards affixed to each kit."
    ))

    # === PAGE 2 ===
    add_heading_paragraph(doc, "3. Machine Safety and Lockout/Tagout (LOTO)", level=1)

    add_body_paragraph(doc, (
        "All machinery with moving parts, electrical hazards, or stored energy must be properly "
        "locked out and tagged out before any maintenance, cleaning, or adjustment work begins. "
        "Only trained and authorized personnel may perform LOTO procedures. Each authorized "
        "employee is issued a personal padlock and tag set that must not be shared."
    ))

    add_body_paragraph(doc, (
        "Before initiating LOTO, the authorized employee must identify all energy sources "
        "associated with the equipment, notify affected employees, and follow the machine-specific "
        "LOTO procedure documented in the Equipment Safety Manual. Energy sources include "
        "electrical, hydraulic, pneumatic, thermal, chemical, and gravitational."
    ))

    add_body_paragraph(doc, (
        "Verification of zero energy state is mandatory after applying locks and tags. This "
        "includes attempting to start the equipment using normal controls, testing circuits "
        "with appropriate measuring instruments, and visually confirming that all moving parts "
        "have come to a complete stop. Only after verification may work proceed."
    ))

    add_heading_paragraph(doc, "4. Fire Prevention and Emergency Response", level=1)

    add_body_paragraph(doc, (
        "Fire extinguishers are inspected monthly and serviced annually by a certified vendor. "
        "Employees should know the locations of the nearest fire extinguisher, pull station, and "
        "emergency exit relative to their work area. The PASS method (Pull, Aim, Squeeze, Sweep) "
        "should be used when operating a portable fire extinguisher."
    ))

    add_body_paragraph(doc, (
        "Evacuation routes are posted in all hallways and common areas. During a fire alarm, "
        "employees must immediately cease all work, close but do not lock doors behind them, "
        "and proceed to the designated assembly point. Roll call will be conducted by department "
        "supervisors. Do not re-enter the building until the all-clear is given by the fire marshal."
    ))

    add_body_paragraph(doc, (
        "Emergency response teams (ERTs) are stationed on each floor of the facility. ERT members "
        "receive advanced first aid, CPR, AED, and fire suppression training. They wear orange "
        "vests during emergency events and are responsible for sweeping their assigned zones "
        "to ensure complete evacuation."
    ))

    add_heading_paragraph(doc, "5. Ergonomics and Office Safety", level=1)

    add_body_paragraph(doc, (
        "Workstation ergonomics assessments are available upon request through the EHS portal. "
        "Recommended practices include positioning monitors at arm's length with the top of the "
        "screen at or below eye level, keeping feet flat on the floor or on a footrest, and "
        "taking micro-breaks every 30 minutes to reduce repetitive strain."
    ))

    add_body_paragraph(doc, (
        "Slip, trip, and fall hazards are the most common cause of workplace injuries in office "
        "environments. Keep walkways clear of cables, boxes, and personal items. Report loose "
        "carpet tiles, wet floors, or damaged stair treads immediately to Facilities Management "
        "at extension 3200."
    ))

    # === PAGE 3 ===
    add_heading_paragraph(doc, "6. Personal Protective Equipment (PPE)", level=1)

    add_body_paragraph(doc, (
        "The selection of appropriate PPE is determined through a Job Hazard Analysis (JHA) "
        "conducted by the department safety officer in coordination with EHS. PPE requirements "
        "vary by task and work area. Common categories include eye and face protection, hearing "
        "protection, respiratory protection, hand protection, and foot protection."
    ))

    # The WARNING paragraph - same formatting as all others, no special treatment
    add_body_paragraph(doc, (
        "WARNING: Always wear protective equipment when entering designated hazard zones. "
        "This includes safety glasses with side shields (ANSI Z87.1 rated), steel-toed boots "
        "(ASTM F2413 compliant), high-visibility vests, hard hats where overhead hazards exist, "
        "and chemical-resistant gloves when handling corrosive or toxic substances. Failure to "
        "wear required PPE may result in immediate removal from the work area and disciplinary "
        "action up to and including termination."
    ))

    add_body_paragraph(doc, (
        "PPE must be inspected before each use for signs of damage, wear, or contamination. "
        "Damaged equipment must be removed from service immediately and replaced. The company "
        "provides all required PPE at no cost to employees. Replacement requests can be submitted "
        "through the supply portal or directly to the department safety officer."
    ))

    add_heading_paragraph(doc, "7. Confined Space Entry", level=1)

    add_body_paragraph(doc, (
        "Entry into confined spaces such as tanks, silos, vessels, and underground vaults requires "
        "a valid Confined Space Entry Permit signed by the area supervisor and a qualified attendant. "
        "Atmospheric testing for oxygen levels, combustible gases, and toxic substances must be "
        "performed before entry and continuously monitored throughout the work period."
    ))

    add_body_paragraph(doc, (
        "A trained attendant must remain at the entry point at all times during confined space "
        "work. The attendant monitors conditions, maintains communication with entrants, and "
        "initiates rescue procedures if necessary. Under no circumstances should the attendant "
        "enter the confined space to attempt a rescue without proper training and equipment."
    ))

    add_heading_paragraph(doc, "8. Electrical Safety", level=1)

    add_body_paragraph(doc, (
        "Only qualified electricians may work on electrical systems rated above 50 volts. "
        "All electrical panels must have clear access with a minimum 36-inch clearance. "
        "Extension cords are for temporary use only and must not be used as permanent wiring. "
        "Ground Fault Circuit Interrupters (GFCIs) are required in all wet locations."
    ))

    add_body_paragraph(doc, (
        "Electrical safety training is required for all employees who may encounter electrical "
        "hazards in their work. This includes understanding arc flash boundaries, using "
        "appropriate voltage-rated tools and PPE, and recognizing signs of electrical faults "
        "such as burning smells, discolored outlets, or flickering lights."
    ))

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
