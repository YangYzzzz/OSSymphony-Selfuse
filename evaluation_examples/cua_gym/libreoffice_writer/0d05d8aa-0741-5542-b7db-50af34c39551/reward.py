"""
Reward Script: Apply paragraph borders to create a callout box around a warning paragraph
Task ID: writer_rd_047
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35) - All 4 borders present with red (#CC0000) color and 2pt width
  Component 2 (0.15) - Border padding (spacing to contents) ~0.5cm on all sides
  Component 3 (0.25) - Paragraph background color is light yellow (#FFFDE6)
  Component 4 (0.25) - 'WARNING:' text is bold and red (#CC0000)
"""

import os
from math import sqrt
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_047'


def color_distance_hex(c1, c2):
    """Compute Euclidean distance between two hex color strings (e.g. 'CC0000')."""
    r1, g1, b1 = int(c1[0:2], 16), int(c1[2:4], 16), int(c1[4:6], 16)
    r2, g2, b2 = int(c2[0:2], 16), int(c2[2:4], 16), int(c2[4:6], 16)
    return sqrt((r1-r2)**2 + (g1-g2)**2 + (b1-b2)**2)


def persist_app_state(domain):
    """Save any unsaved LibreOffice changes."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print("PERSIST: ctrl+s sent for %s" % domain)
        except Exception as e:
            print("PERSIST_WARN: save hook failed: %s" % e)


def find_warning_paragraph(doc):
    """Find the paragraph containing 'WARNING:' text."""
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith('WARNING:'):
            return i, para
    # Fallback: any paragraph containing WARNING
    for i, para in enumerate(doc.paragraphs):
        if 'WARNING' in para.text:
            return i, para
    return None, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    # Find the WARNING paragraph
    idx, warn_para = find_warning_paragraph(doc)
    if warn_para is None:
        print("CRITICAL: No WARNING paragraph found in document")
        print("REWARD: 0.0")
        return 0.0

    print("Found WARNING paragraph at index %d: %s..." % (idx, warn_para.text[:80]))

    pPr = warn_para._element.find(qn('w:pPr'))

    # Component 1: All 4 borders present with red (#CC0000) color and ~2pt width (0.35 pts)
    try:
        pBdr = pPr.find(qn('w:pBdr')) if pPr is not None else None
        if pBdr is None:
            print("FAIL: Component 1 — No paragraph borders found")
        else:
            sides_ok = 0
            for side in ['top', 'left', 'bottom', 'right']:
                el = pBdr.find(qn('w:' + side))
                if el is not None:
                    w_val = el.get(qn('w:val'))
                    w_sz = el.get(qn('w:sz'))
                    w_color = el.get(qn('w:color'))

                    # Check border exists and is visible (not 'none' or 'nil')
                    if w_val in ('none', 'nil', None):
                        print("FAIL: Component 1 — %s border has val=%s (not visible)" % (side, w_val))
                        continue

                    # Check color is close to CC0000 (red)
                    color_ok = False
                    if w_color is not None:
                        try:
                            dist = color_distance_hex(w_color.upper(), 'CC0000')
                            color_ok = dist < 50  # tolerance for similar reds
                        except Exception:
                            pass

                    # Check size >= 12 eighth-pts (1.5pt) — task says 2pt = 16 eighth-pts
                    sz_ok = False
                    if w_sz is not None:
                        try:
                            sz_val = int(w_sz)
                            sz_ok = sz_val >= 12  # at least ~1.5pt
                        except ValueError:
                            pass

                    if color_ok and sz_ok:
                        sides_ok += 1
                        print("PASS: Component 1 — %s border: val=%s, sz=%s, color=%s" % (side, w_val, w_sz, w_color))
                    else:
                        print("FAIL: Component 1 — %s border: color_ok=%s, sz_ok=%s (val=%s, sz=%s, color=%s)" %
                              (side, color_ok, sz_ok, w_val, w_sz, w_color))
                else:
                    print("FAIL: Component 1 — %s border element missing" % side)

            if sides_ok == 4:
                print("PASS: Component 1 — All 4 borders correct (0.35 pts)")
                total_score += 0.35
            elif sides_ok > 0:
                partial = 0.35 * (sides_ok / 4.0)
                print("PARTIAL: Component 1 — %d/4 borders correct (%.2f pts)" % (sides_ok, partial))
                total_score += partial
            else:
                print("FAIL: Component 1 — No borders matched criteria")
    except Exception as e:
        print("ERROR: Component 1 — %s" % e)

    # Component 2: Border padding ~0.5cm (space attribute) on all sides (0.15 pts)
    # 0.5cm ≈ 14 eighth-points in OOXML. Accept range 10-20 for tolerance.
    try:
        pBdr = pPr.find(qn('w:pBdr')) if pPr is not None else None
        if pBdr is None:
            print("FAIL: Component 2 — No borders, cannot check padding")
        else:
            padding_ok = 0
            for side in ['top', 'left', 'bottom', 'right']:
                el = pBdr.find(qn('w:' + side))
                if el is not None:
                    w_space = el.get(qn('w:space'))
                    if w_space is not None:
                        try:
                            space_val = int(w_space)
                            if 8 <= space_val <= 22:  # ~0.3cm to ~0.8cm tolerance
                                padding_ok += 1
                                print("PASS: Component 2 — %s padding: space=%s" % (side, w_space))
                            else:
                                print("FAIL: Component 2 — %s padding: space=%s (expected 10-20)" % (side, w_space))
                        except ValueError:
                            print("FAIL: Component 2 — %s padding: invalid space=%s" % (side, w_space))
                    else:
                        print("FAIL: Component 2 — %s padding: no space attribute" % side)

            if padding_ok == 4:
                print("PASS: Component 2 — All 4 sides have correct padding (0.15 pts)")
                total_score += 0.15
            elif padding_ok > 0:
                partial = 0.15 * (padding_ok / 4.0)
                print("PARTIAL: Component 2 — %d/4 sides padding OK (%.2f pts)" % (padding_ok, partial))
                total_score += partial
            else:
                print("FAIL: Component 2 — No padding matched criteria")
    except Exception as e:
        print("ERROR: Component 2 — %s" % e)

    # Component 3: Paragraph background color is light yellow #FFFDE6 (0.25 pts)
    try:
        shd = pPr.find(qn('w:shd')) if pPr is not None else None
        if shd is None:
            print("FAIL: Component 3 — No paragraph shading found")
        else:
            fill_color = shd.get(qn('w:fill'))
            if fill_color is not None:
                try:
                    dist = color_distance_hex(fill_color.upper(), 'FFFDE6')
                    if dist < 30:  # close to FFFDE6
                        print("PASS: Component 3 — Background fill=%s (0.25 pts)" % fill_color)
                        total_score += 0.25
                    else:
                        print("FAIL: Component 3 — Background fill=%s, distance=%.1f from FFFDE6" % (fill_color, dist))
                except Exception as ex:
                    print("FAIL: Component 3 — Cannot parse fill color %s: %s" % (fill_color, ex))
            else:
                print("FAIL: Component 3 — No fill attribute in shading element")
    except Exception as e:
        print("ERROR: Component 3 — %s" % e)

    # Component 4: 'WARNING:' text is bold and red (#CC0000) (0.25 pts)
    try:
        runs = warn_para.runs
        if not runs:
            print("FAIL: Component 4 — No runs found in WARNING paragraph")
        else:
            # Find the run(s) containing 'WARNING:'
            warning_bold = False
            warning_red = False
            for run in runs:
                if 'WARNING' in run.text:
                    # Check bold
                    if run.bold is True:
                        warning_bold = True
                    # Check red color
                    if run.font.color.rgb is not None:
                        try:
                            rgb_str = str(run.font.color.rgb).upper()
                            dist = color_distance_hex(rgb_str, 'CC0000')
                            if dist < 50:
                                warning_red = True
                        except Exception:
                            pass
                    print("  Run with WARNING: text=%s, bold=%s, color=%s" % (repr(run.text[:30]), run.bold, run.font.color.rgb))
                    break  # Only check first run containing WARNING

            if warning_bold and warning_red:
                print("PASS: Component 4 — 'WARNING:' is bold and red (0.25 pts)")
                total_score += 0.25
            elif warning_bold or warning_red:
                partial = 0.125
                print("PARTIAL: Component 4 — bold=%s, red=%s (%.3f pts)" % (warning_bold, warning_red, partial))
                total_score += partial
            else:
                print("FAIL: Component 4 — 'WARNING:' not bold (%s) or red (%s)" % (warning_bold, warning_red))
    except Exception as e:
        print("ERROR: Component 4 — %s" % e)

    final_score = min(round(total_score, 2), 1.0)
    print("\nScore: %.2f/1.0" % total_score)
    print("REWARD: %.1f" % final_score)
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = os.path.join(WORKDIR, TASK_ID + '.docx')
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
