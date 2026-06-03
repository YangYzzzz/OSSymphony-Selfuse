"""
Initial Setup: Insert date field in header of a contract
Task ID: writer_legal_017
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_017'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # -- Header: PRIVILEGED AND CONFIDENTIAL, left-aligned, NO date --
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = hp.add_run("PRIVILEGED AND CONFIDENTIAL")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # -- Footer with page number --
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fp.add_run("Page ").font.size = Pt(9)
    from docx.oxml.ns import qn
    r1 = fp.add_run()
    r1.font.size = Pt(9)
    r1._element.append(r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}))
    r2 = fp.add_run()
    r2.font.size = Pt(9)
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    r3.font.size = Pt(9)
    r3._element.append(r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}))

    # -- Document Title --
    title = doc.add_heading("PROFESSIONAL SERVICES AGREEMENT", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -- Parties --
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    run = intro.add_run("This Professional Services Agreement")
    run.bold = True
    intro.add_run(' (the "Agreement") is entered into as of March 15, 2026 (the "Effective Date"), by and between:')

    party1 = doc.add_paragraph()
    party1.paragraph_format.left_indent = Inches(0.5)
    party1.paragraph_format.space_after = Pt(4)
    r = party1.add_run("Westbrook Analytics, Inc.")
    r.bold = True
    party1.add_run(", a Delaware corporation with principal offices at 2400 Market Street, Suite 1200, San Francisco, CA 94114 (\"Company\"); and")

    party2 = doc.add_paragraph()
    party2.paragraph_format.left_indent = Inches(0.5)
    party2.paragraph_format.space_after = Pt(12)
    r = party2.add_run("Meridian Consulting Group, LLC")
    r.bold = True
    party2.add_run(", a New York limited liability company with offices at 350 Fifth Avenue, 42nd Floor, New York, NY 10118 (\"Consultant\").")

    doc.add_paragraph("WHEREAS, Company desires to retain Consultant to provide certain professional consulting services; and")
    doc.add_paragraph("WHEREAS, Consultant has the requisite expertise, experience, and personnel to perform such services;")
    doc.add_paragraph("NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth herein, the parties agree as follows:")

    # -- Article 1: Scope of Services --
    doc.add_heading("ARTICLE 1: SCOPE OF SERVICES", level=1)

    p = doc.add_paragraph()
    r = p.add_run("1.1 Services. ")
    r.bold = True
    p.add_run("Consultant shall provide the professional services described in Exhibit A attached hereto and incorporated herein by reference (the \"Services\"). The Services shall include, but not be limited to, strategic business analysis, market research, financial modeling, and preparation of deliverables as specified therein.")

    p = doc.add_paragraph()
    r = p.add_run("1.2 Standard of Performance. ")
    r.bold = True
    p.add_run("Consultant shall perform the Services in a professional and workmanlike manner, consistent with generally accepted industry standards and practices. Consultant represents that it possesses the necessary qualifications, experience, and expertise to perform the Services.")

    p = doc.add_paragraph()
    r = p.add_run("1.3 Personnel. ")
    r.bold = True
    p.add_run("Consultant shall assign the following key personnel to this engagement: Dr. Rachel Thornton (Lead Analyst), James K. Whitfield (Senior Consultant), and Anita Ramirez (Project Coordinator). Consultant shall not replace key personnel without the prior written consent of Company.")

    # -- Article 2: Compensation --
    doc.add_heading("ARTICLE 2: COMPENSATION AND PAYMENT", level=1)

    p = doc.add_paragraph()
    r = p.add_run("2.1 Fees. ")
    r.bold = True
    p.add_run("Company shall pay Consultant a fixed fee of $187,500.00 for the Services, payable in three equal installments of $62,500.00 each, due upon completion of the milestones set forth in Exhibit B.")

    p = doc.add_paragraph()
    r = p.add_run("2.2 Expenses. ")
    r.bold = True
    p.add_run("Company shall reimburse Consultant for reasonable and pre-approved out-of-pocket expenses incurred in connection with the performance of Services. Total reimbursable expenses shall not exceed $15,000.00 without prior written approval from Company.")

    p = doc.add_paragraph()
    r = p.add_run("2.3 Invoicing. ")
    r.bold = True
    p.add_run("Consultant shall submit invoices within fifteen (15) business days following achievement of each milestone. Company shall pay undisputed invoices within thirty (30) days of receipt. Late payments shall accrue interest at the rate of 1.5% per month.")

    # -- Article 3: Term and Termination --
    doc.add_heading("ARTICLE 3: TERM AND TERMINATION", level=1)

    p = doc.add_paragraph()
    r = p.add_run("3.1 Term. ")
    r.bold = True
    p.add_run("This Agreement shall commence on the Effective Date and shall continue for a period of twelve (12) months, unless earlier terminated in accordance with this Article 3.")

    p = doc.add_paragraph()
    r = p.add_run("3.2 Termination for Convenience. ")
    r.bold = True
    p.add_run("Either party may terminate this Agreement upon thirty (30) days' prior written notice to the other party. In the event of such termination, Company shall pay Consultant for all Services satisfactorily performed through the date of termination.")

    p = doc.add_paragraph()
    r = p.add_run("3.3 Termination for Cause. ")
    r.bold = True
    p.add_run("Either party may terminate this Agreement immediately upon written notice if the other party: (a) materially breaches any provision of this Agreement and fails to cure such breach within fifteen (15) days after receiving written notice thereof; or (b) becomes insolvent, files for bankruptcy, or has a receiver appointed for its assets.")

    # -- Article 4: Confidentiality --
    doc.add_heading("ARTICLE 4: CONFIDENTIALITY", level=1)

    p = doc.add_paragraph()
    r = p.add_run("4.1 Confidential Information. ")
    r.bold = True
    p.add_run("Each party acknowledges that in the course of performing this Agreement, it may receive or have access to confidential and proprietary information of the other party, including but not limited to trade secrets, business plans, financial data, customer lists, technical specifications, and marketing strategies (\"Confidential Information\").")

    p = doc.add_paragraph()
    r = p.add_run("4.2 Non-Disclosure. ")
    r.bold = True
    p.add_run("The receiving party shall: (a) hold all Confidential Information in strict confidence; (b) not disclose Confidential Information to any third party without the prior written consent of the disclosing party; and (c) use Confidential Information solely for the purposes of this Agreement.")

    # -- Signature Block --
    doc.add_paragraph("")  # spacer
    doc.add_heading("IN WITNESS WHEREOF", level=2)
    doc.add_paragraph("The parties have executed this Agreement as of the date first written above.")

    # Signature table
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "WESTBROOK ANALYTICS, INC."
    table.cell(0, 1).text = "MERIDIAN CONSULTING GROUP, LLC"
    table.cell(1, 0).text = "By: _________________________"
    table.cell(1, 1).text = "By: _________________________"
    table.cell(2, 0).text = "Name: Victoria S. Hartwell"
    table.cell(2, 1).text = "Name: David R. Patterson"
    table.cell(3, 0).text = "Title: Chief Executive Officer"
    table.cell(3, 1).text = "Title: Managing Partner"

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
