"""
Reward Script: Insert today's date as a field in the header, right-aligned
Task ID: writer_legal_017
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): DATE field code exists in header
  Component 2 (0.3): Right-alignment mechanism for date (right tab stop or right-aligned para)
  Component 3 (0.3): Complete field structure (begin/separate/end) with confidential text preserved
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_017'


def persist_app_state(domain):
    """Save any unsaved LibreOffice changes before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Get header from first section
    try:
        header = doc.sections[0].header
        header_xml = header._element.xml
    except Exception as e:
        print(f"CRITICAL: Cannot access header: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: DATE field code exists in header (0.4 points)
    # This checks that a DATE field (not static text) was inserted.
    # FAILS on initial (no field codes), PASSES on golden.
    try:
        # Look for instrText containing DATE in the header XML
        has_date_field = bool(re.search(r'instrText[^>]*>.*?DATE', header_xml, re.DOTALL))
        if has_date_field:
            print(f"PASS: Component 1 — DATE field code found in header (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — No DATE field code found in header XML")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Right-alignment mechanism for date (0.3 points)
    # The date should be right-aligned. This can be via:
    #   (a) A right-aligned tab stop in the header paragraph, OR
    #   (b) A separate right-aligned paragraph
    # FAILS on initial (no right tab stop, no right-aligned para), PASSES on golden.
    try:
        has_right_alignment = False

        for para in header.paragraphs:
            para_xml = para._element.xml
            # Check for right tab stop: <w:tab w:val="right" .../>
            if re.search(r'<w:tab[^>]*w:val="right"', para_xml):
                has_right_alignment = True
                break
            # Check for right-aligned paragraph with field code
            if re.search(r'<w:jc\s+w:val="right"', para_xml) and 'instrText' in para_xml:
                has_right_alignment = True
                break

        if has_right_alignment:
            print(f"PASS: Component 2 — Right-alignment mechanism found for date (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — No right-alignment mechanism found for date in header")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Complete field structure AND confidential text preserved (0.3 points)
    # The field must have proper begin/separate/end structure AND the original
    # "PRIVILEGED AND CONFIDENTIAL" text must still be present.
    # FAILS on initial (no fldChar at all), PASSES on golden.
    try:
        has_complete_field = False
        has_confidential = False

        # Check for complete field structure: fldCharType="begin" ... "separate" ... "end"
        fld_begin = 'fldCharType="begin"' in header_xml
        fld_separate = 'fldCharType="separate"' in header_xml
        fld_end = 'fldCharType="end"' in header_xml
        has_complete_field = fld_begin and fld_separate and fld_end

        # Check confidential text preserved
        full_header_text = ' '.join(p.text for p in header.paragraphs)
        has_confidential = 'PRIVILEGED AND CONFIDENTIAL' in full_header_text

        if has_complete_field and has_confidential:
            print(f"PASS: Component 3 — Complete field structure and confidential text preserved (0.3 pts)")
            total_score += 0.3
        else:
            if not has_complete_field:
                print(f"FAIL: Component 3 — Incomplete field structure (begin={fld_begin}, separate={fld_separate}, end={fld_end})")
            if not has_confidential:
                print(f"FAIL: Component 3 — 'PRIVILEGED AND CONFIDENTIAL' text missing from header (found: {repr(full_header_text[:100])})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
