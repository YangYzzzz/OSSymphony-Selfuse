"""
Initial Setup: Create a multi-chapter business document with empty headers/footers
Task ID: writer_biz_071
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_071'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Ensure header and footer are empty but enabled
    header = section.header
    header.is_linked_to_previous = False
    # Clear any default content
    for p in header.paragraphs:
        p.text = ""

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.text = ""

    # --- Chapter 1: Executive Summary ---
    h1 = doc.add_heading('Executive Summary', level=1)

    doc.add_paragraph(
        'Meridian Solutions Inc. has experienced significant growth over the past fiscal year, '
        'driven primarily by expansion into the Asia-Pacific region and the successful launch of '
        'our CloudBridge platform. Revenue increased by 23% year-over-year, reaching $184.7 million '
        'in Q4 2025 alone.'
    )
    doc.add_paragraph(
        'This strategic review outlines our key achievements, identifies emerging market opportunities, '
        'and presents a comprehensive roadmap for sustainable growth through 2027. Our focus remains '
        'on delivering enterprise-grade solutions while maintaining the agility that has defined our '
        'competitive advantage.'
    )
    doc.add_paragraph(
        'The board of directors has approved an accelerated investment plan targeting three core areas: '
        'artificial intelligence integration, cybersecurity services, and managed cloud infrastructure. '
        'Combined, these initiatives represent a $42 million commitment over the next 18 months.'
    )

    # --- Chapter 2: Market Analysis ---
    doc.add_page_break()
    h2 = doc.add_heading('Market Analysis', level=1)

    doc.add_paragraph(
        'The global enterprise software market is projected to reach $892 billion by 2027, growing at '
        'a compound annual rate of 11.3%. Within this landscape, cloud-native solutions continue to '
        'capture an increasing share, now representing 67% of new enterprise deployments.'
    )
    doc.add_paragraph(
        'Our primary competitors - Apex Digital, NovaTech Systems, and Horizon Group - have each '
        'announced major platform updates in the past quarter. However, independent analyst reviews '
        'from Gartner and Forrester consistently rank Meridian Solutions in the top quadrant for both '
        'innovation and execution capability.'
    )
    doc.add_paragraph(
        'Key market trends affecting our strategy include the rapid adoption of AI-powered analytics, '
        'increasing regulatory requirements around data sovereignty, and the shift toward consumption-based '
        'pricing models. Our product roadmap addresses each of these dynamics with targeted feature releases '
        'scheduled for Q2 and Q3 2026.'
    )
    doc.add_paragraph(
        'Regional analysis shows particularly strong demand in Southeast Asia, where enterprise IT spending '
        'grew by 18.4% in the last fiscal year. Our Singapore office, established in March 2025, has already '
        'secured contracts valued at $12.3 million with major financial institutions.'
    )

    # --- Chapter 3: Product Development ---
    doc.add_page_break()
    h3 = doc.add_heading('Product Development', level=1)

    doc.add_paragraph(
        'The engineering division completed 847 feature requests and resolved 2,341 issues during Q4 2025. '
        'Our CloudBridge 3.0 release introduced multi-tenant architecture, real-time collaboration features, '
        'and a redesigned API gateway that reduced latency by 40%.'
    )
    doc.add_paragraph(
        'Project Aurora, our next-generation AI integration framework, entered beta testing in December 2025 '
        'with 23 enterprise customers participating. Early feedback has been overwhelmingly positive, with '
        'pilot users reporting a 35% reduction in manual data processing tasks.'
    )
    doc.add_paragraph(
        'The development team has grown to 312 engineers across our offices in Austin, Toronto, Singapore, '
        'and Berlin. We have maintained a 94% retention rate among senior engineers, well above the industry '
        'average of 78%. Our technical interview process was redesigned in Q3, resulting in a 28% improvement '
        'in offer acceptance rates.'
    )

    # --- Chapter 4: Financial Performance ---
    doc.add_page_break()
    h4 = doc.add_heading('Financial Performance', level=1)

    doc.add_paragraph(
        'Total revenue for fiscal year 2025 reached $687.2 million, representing a 23% increase over '
        'the previous year. Recurring revenue from subscription services accounted for 72% of total revenue, '
        'up from 64% in FY2024. This shift toward predictable revenue streams strengthens our financial '
        'position and supports long-term investment planning.'
    )
    doc.add_paragraph(
        'Operating margins improved to 18.7%, driven by automation initiatives in customer support and '
        'infrastructure optimization. The finance team implemented a new cost allocation framework that '
        'provides department-level visibility into resource consumption, enabling more precise budget management.'
    )
    doc.add_paragraph(
        'Cash reserves stand at $234 million as of December 31, 2025. The company completed a secondary '
        'offering in September, raising $95 million to fund strategic acquisitions. Two acquisition targets '
        'have been identified and due diligence is underway for DataVault Analytics, a mid-market data '
        'governance platform based in Copenhagen.'
    )

    # --- Chapter 5: Strategic Outlook ---
    doc.add_page_break()
    h5 = doc.add_heading('Strategic Outlook', level=1)

    doc.add_paragraph(
        'Looking ahead to fiscal year 2026, Meridian Solutions is positioned to capitalize on several '
        'converging market trends. Our three-pillar growth strategy focuses on organic product expansion, '
        'strategic acquisitions, and international market penetration.'
    )
    doc.add_paragraph(
        'The leadership team has established ambitious but achievable targets: $850 million in total revenue, '
        '80% recurring revenue ratio, and expansion into three new geographic markets. These targets are '
        'supported by a detailed execution plan that has been reviewed and endorsed by the board.'
    )
    doc.add_paragraph(
        'Risk factors under active monitoring include potential regulatory changes in the European Union '
        'regarding AI governance, currency fluctuations affecting our international revenue streams, and '
        'talent competition in key technology hubs. Mitigation strategies for each risk area are detailed '
        'in Appendix B of this report.'
    )
    doc.add_paragraph(
        'In conclusion, Meridian Solutions Inc. enters 2026 from a position of strength. Our technology '
        'platform is market-leading, our financial foundation is solid, and our team is motivated and '
        'well-equipped to execute on the ambitious goals ahead. We remain committed to delivering '
        'exceptional value to our customers, partners, and shareholders.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
