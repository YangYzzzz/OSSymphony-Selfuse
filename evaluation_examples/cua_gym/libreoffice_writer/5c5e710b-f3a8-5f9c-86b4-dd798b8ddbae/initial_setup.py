"""
Initial Setup: Create a Writer document with procedure steps as a flat numbered list.
Task ID: writer_tech_033
Domain: libreoffice_writer

All procedure steps (main and sub) are plain numbered paragraphs at the same level,
with no outline numbering or indentation hierarchy.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_033'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading("Network Infrastructure Upgrade Plan", level=0)

    # --- Introduction ---
    doc.add_heading("Introduction", level=1)
    intro = doc.add_paragraph(
        "This document outlines the procedure for upgrading the corporate network "
        "infrastructure from the legacy Cat5e backbone to a modern Cat6a/fiber hybrid "
        "topology. The upgrade is scheduled to occur over the weekend of March 22-23, 2025, "
        "to minimize disruption to business operations."
    )

    doc.add_paragraph(
        "All team members involved in the upgrade should review this procedure carefully "
        "and confirm their availability with the project lead, Sarah Chen, by March 15, 2025."
    )

    # --- Scope ---
    doc.add_heading("Scope", level=1)
    doc.add_paragraph(
        "The upgrade covers Building A floors 3-7 and Building B floors 1-4. "
        "The server room on Building A floor 2 will be handled in a separate phase. "
        "Approximately 340 network drops will be replaced and 12 new access points installed."
    )

    # --- Procedure Section ---
    doc.add_heading("Procedure", level=1)

    doc.add_paragraph(
        "Follow the steps below in order. Each main step must be completed and verified "
        "before proceeding to the next."
    )

    # All 9 items as flat numbered paragraphs at the same level using "List Number" style.
    # This means NO outline numbering, NO indentation difference between main and sub steps.
    # They all appear as: 1. 2. 3. 4. 5. 6. 7. 8. 9.

    steps = [
        "Disconnect all active connections from the legacy switches and document the port mapping for each floor segment.",
        "Verify that the backup communication channel via the 4G failover gateway is operational and accessible to all departments.",
        "Label each cable run at both the patch panel end and the wall outlet end using the naming convention BLD-FLOOR-ROOM-PORT.",
        "Remove the legacy Cat5e patch cables from the main distribution frame in the server room on each floor.",
        "Route the new Cat6a cables through the existing conduit pathways, ensuring minimum bend radius of 25mm is maintained.",
        "Terminate all new cable runs at the Keystone jacks using T568B wiring standard and verify with a continuity tester.",
        "Install and configure the new managed switches in each IDF closet according to the VLAN assignment spreadsheet.",
        "Connect the uplink fiber from each IDF closet to the MDF using the pre-terminated LC-LC OM4 multimode patch cables.",
        "Run the comprehensive network test suite from the monitoring workstation to validate throughput, latency, and packet loss across all segments.",
    ]

    for step_text in steps:
        para = doc.add_paragraph(step_text, style='List Number')

    # --- Post-Procedure Notes ---
    doc.add_heading("Verification and Sign-Off", level=1)
    doc.add_paragraph(
        "After completing all steps, the network operations team must run a full "
        "connectivity audit using the Nagios monitoring platform. Results should be "
        "documented in the project tracker and submitted to Marcus Johnson for final "
        "approval no later than March 24, 2025."
    )

    doc.add_paragraph(
        "Any failed test points must be escalated to the cabling contractor within "
        "24 hours for remediation under the warranty agreement."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
