"""
Reward Script: Multi-level numbered list with outline numbering
Task ID: writer_tech_033
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Sub-steps at level 1 (6 paragraphs at ilvl=1)
  Component 2 (0.30): Hierarchical structure - main steps ilvl=0, sub-steps ilvl=1, shared numId
  Component 3 (0.25): Correct grouping pattern (main-sub-sub repeating)
  Component 4 (0.15): Outline numbering format (%1.%2. pattern)
"""

import os
import sys

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_033'

# Expected procedure paragraph texts (unique substrings for matching)
MAIN_STEP_TEXTS = [
    "Disconnect all active connections from the legacy switches",
    "Remove the legacy Cat5e patch cables from the main distribution frame",
    "Install and configure the new managed switches in each IDF closet",
]

SUB_STEP_TEXTS = [
    "Verify that the backup communication channel via the 4G failover gateway",
    "Label each cable run at both the patch panel end",
    "Route the new Cat6a cables through the existing conduit pathways",
    "Terminate all new cable runs at the Keystone jacks",
    "Connect the uplink fiber from each IDF closet to the MDF",
    "Run the comprehensive network test suite from the monitoring workstation",
]

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def persist_app_state(domain):
    """Save any unsaved changes in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def get_para_numbering(para):
    """
    Extract numId and ilvl from a paragraph's numPr element.
    Returns (numId, ilvl) or (None, None) if no numbering.
    """
    ns = {'w': WNS}
    numPr = para._element.find('.//w:numPr', ns)
    if numPr is None:
        return None, None
    numId_el = numPr.find('w:numId', ns)
    ilvl_el = numPr.find('w:ilvl', ns)
    numId = numId_el.get(f'{{{WNS}}}val') if numId_el is not None else None
    ilvl = ilvl_el.get(f'{{{WNS}}}val') if ilvl_el is not None else None
    return numId, ilvl


def find_para_by_text(doc, substring):
    """Find a paragraph whose text starts with the given substring."""
    for para in doc.paragraphs:
        if para.text.strip().startswith(substring[:40]):
            return para
    return None


def get_abstract_num_for_numid(doc, numId):
    """Get the abstractNum element for a given numId."""
    ns = {'w': WNS}
    try:
        numbering_elem = doc.part.numbering_part.element
        for num in numbering_elem.findall('.//w:num', ns):
            nid = num.get(f'{{{WNS}}}numId')
            if nid == str(numId):
                absRef = num.find('w:abstractNumId', ns)
                if absRef is not None:
                    absId = absRef.get(f'{{{WNS}}}val')
                    for absNum in numbering_elem.findall('.//w:abstractNum', ns):
                        if absNum.get(f'{{{WNS}}}abstractNumId') == absId:
                            return absNum
    except Exception:
        pass
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': WNS}
    sub_correct = 0  # track across components

    # Component 1: Sub-steps exist at level 1 (0.30 points)
    # At least some sub-step paragraphs must have ilvl=1 in numPr
    # This is the fundamental multi-level change: initial has all at ilvl=0
    try:
        sub_correct = 0
        for text_prefix in SUB_STEP_TEXTS:
            para = find_para_by_text(doc, text_prefix)
            if para is not None:
                numId, ilvl = get_para_numbering(para)
                if numId is not None and ilvl is not None and int(ilvl) == 1:
                    sub_correct += 1
                    print(f"  Sub-step OK: ilvl=1, text starts with '{text_prefix[:50]}...'")
                else:
                    print(f"  Sub-step FAIL: numId={numId}, ilvl={ilvl}, text starts with '{text_prefix[:50]}...'")
            else:
                print(f"  Sub-step MISSING: Could not find paragraph starting with '{text_prefix[:50]}...'")

        if sub_correct == 6:
            print(f"PASS: Component 1 — All 6 sub-steps at level 1 (0.30 pts)")
            total_score += 0.30
        elif sub_correct > 0:
            partial = 0.30 * (sub_correct / 6.0)
            print(f"PARTIAL: Component 1 — {sub_correct}/6 sub-steps at level 1 ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No sub-steps at level 1")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Main steps at level 0 AND sub-steps at level 1 under same numId (0.30 points)
    # Verifies the hierarchical structure: main steps ilvl=0 share a numId with sub-steps at ilvl=1
    try:
        main_correct = 0
        shared_numid_correct = 0
        main_numIds = []
        sub_numIds = []

        for text_prefix in MAIN_STEP_TEXTS:
            para = find_para_by_text(doc, text_prefix)
            if para is not None:
                numId, ilvl = get_para_numbering(para)
                if numId is not None and ilvl is not None and int(ilvl) == 0:
                    main_correct += 1
                    main_numIds.append(numId)
                    print(f"  Main step OK: numId={numId} ilvl=0, '{text_prefix[:50]}...'")
                else:
                    print(f"  Main step FAIL: numId={numId}, ilvl={ilvl}, '{text_prefix[:50]}...'")

        for text_prefix in SUB_STEP_TEXTS:
            para = find_para_by_text(doc, text_prefix)
            if para is not None:
                numId, ilvl = get_para_numbering(para)
                if numId is not None:
                    sub_numIds.append(numId)

        # Check that main steps and sub-steps share a common numId (same list)
        main_set = set(main_numIds)
        sub_set = set(sub_numIds)
        shared = main_set & sub_set
        all_share_numid = len(shared) > 0 and main_set == shared and sub_set == shared

        # Only award points if sub-steps are also at level 1 (prevents scoring flat lists)
        if main_correct == 3 and all_share_numid and sub_correct == 6:
            print(f"PASS: Component 2 — Hierarchical structure correct: 3 main (ilvl=0) + 6 sub (ilvl=1) share numId (0.30 pts)")
            total_score += 0.30
        elif main_correct == 3 and sub_correct > 0:
            # Partial: structure mostly right but numIds may differ
            partial = 0.20
            print(f"PARTIAL: Component 2 — Main steps OK but numId sharing incomplete ({partial} pts)")
            total_score += partial
        elif main_correct > 0 and sub_correct > 0:
            partial = 0.10
            print(f"PARTIAL: Component 2 — Some main steps at level 0, partial sub-steps ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Hierarchical structure not found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Correct grouping pattern (0.25 points)
    # Verify the order: main-sub-sub-main-sub-sub-main-sub-sub
    # Each main step paragraph is followed by exactly 2 sub-step paragraphs
    try:
        # Get ilvl for all procedure paragraphs (P8-P16 in the Procedure section)
        procedure_levels = []
        in_procedure = False
        for para in doc.paragraphs:
            if para.text.strip() == 'Procedure' and para.style and 'Heading' in para.style.name:
                in_procedure = True
                continue
            if in_procedure:
                # Stop at the next heading
                if para.style and 'Heading' in para.style.name:
                    break
                numId, ilvl = get_para_numbering(para)
                if numId is not None and ilvl is not None:
                    procedure_levels.append(int(ilvl))

        expected_pattern = [0, 1, 1, 0, 1, 1, 0, 1, 1]
        print(f"  Procedure numbering levels: {procedure_levels}")
        print(f"  Expected pattern:           {expected_pattern}")

        if procedure_levels == expected_pattern:
            print(f"PASS: Component 3 — Correct grouping pattern (0.25 pts)")
            total_score += 0.25
        elif len(procedure_levels) == 9 and procedure_levels.count(0) == 3 and procedure_levels.count(1) == 6:
            # Right counts but possibly wrong order
            print(f"PARTIAL: Component 3 — Right level counts but wrong order (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 — Grouping pattern does not match expected")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Outline numbering format (0.15 points)
    # The numbering definition for the list should use outline-style format
    # (level 1 text pattern includes parent level, e.g., %1.%2.)
    try:
        outline_format = False
        # Find any numbered paragraph to get numId
        sample_numId = None
        for para in doc.paragraphs:
            nid, ilv = get_para_numbering(para)
            if nid is not None:
                sample_numId = nid
                break

        if sample_numId is not None:
            absNum = get_abstract_num_for_numid(doc, sample_numId)
            if absNum is not None:
                # Check level 1's lvlText for outline pattern (should contain %1.%2)
                for lvl in absNum.findall('w:lvl', ns):
                    ilvl_val = lvl.get(f'{{{WNS}}}ilvl')
                    if ilvl_val == '1':
                        lvlText_el = lvl.find('w:lvlText', ns)
                        numFmt_el = lvl.find('w:numFmt', ns)
                        if lvlText_el is not None:
                            lvlText = lvlText_el.get(f'{{{WNS}}}val', '')
                            numFmt = numFmt_el.get(f'{{{WNS}}}val', '') if numFmt_el is not None else ''
                            print(f"  Level 1 format: lvlText='{lvlText}' numFmt='{numFmt}'")
                            # Outline numbering: level 1 text includes %1 (parent reference)
                            if '%1' in lvlText and '%2' in lvlText and numFmt == 'decimal':
                                outline_format = True
                        break
                # Also check level 0 uses decimal
                for lvl in absNum.findall('w:lvl', ns):
                    ilvl_val = lvl.get(f'{{{WNS}}}ilvl')
                    if ilvl_val == '0':
                        numFmt_el = lvl.find('w:numFmt', ns)
                        if numFmt_el is not None:
                            fmt = numFmt_el.get(f'{{{WNS}}}val', '')
                            print(f"  Level 0 format: numFmt='{fmt}'")
                        break

        if outline_format:
            print(f"PASS: Component 4 — Outline numbering format (%1.%2. pattern) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 4 — Numbering does not use outline format")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
