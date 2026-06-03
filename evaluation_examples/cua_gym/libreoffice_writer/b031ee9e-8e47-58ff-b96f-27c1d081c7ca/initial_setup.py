"""
Initial Setup: Academic document with Appendix A following References without page style change.
Task ID: writer_acad_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_041'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard academic margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ---- Title ----
    title = doc.add_heading('The Impact of Remote Work on Employee Productivity:\nA Mixed-Methods Study', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Author info
    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('Dr. Elena Vasquez, Department of Organizational Behavior\n')
    run.font.size = Pt(12)
    run = author.add_run('University of Western Ontario\n')
    run.font.size = Pt(11)
    run = author.add_run('Published: March 2025')
    run.font.size = Pt(11)
    run.italic = True

    # ---- Abstract ----
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This study examines the relationship between remote work arrangements and employee '
        'productivity across 14 mid-size technology companies in North America. Using a convergent '
        'mixed-methods design, we collected survey data from 1,247 knowledge workers and conducted '
        'semi-structured interviews with 38 team leads over a 12-month period (January 2024 to '
        'December 2024). Quantitative results indicate a statistically significant increase in '
        'self-reported productivity (M = 4.12, SD = 0.87) among fully remote workers compared to '
        'hybrid (M = 3.78, SD = 0.93) and in-office (M = 3.45, SD = 1.02) counterparts, '
        'F(2, 1244) = 18.67, p < .001. Qualitative analysis revealed three emergent themes: '
        'autonomy-driven focus, communication overhead, and boundary management challenges. '
        'Findings suggest that remote work benefits are moderated by organizational support '
        'structures and individual self-regulation capacity.'
    )

    # ---- Introduction ----
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'The global shift toward remote and hybrid work arrangements, accelerated by the '
        'COVID-19 pandemic, has fundamentally altered how organizations conceptualize productivity '
        'and workplace design (Bloom et al., 2023). While early pandemic-era studies focused on '
        'immediate transitions, the field now requires longitudinal evidence on sustained remote '
        'work impacts (Choudhury, 2022). This study addresses that gap by examining productivity '
        'outcomes over a full calendar year across multiple organizations.'
    )
    doc.add_paragraph(
        'Prior research has yielded mixed findings. Some studies report productivity gains of '
        '13-24% among remote workers (Bloom et al., 2015; Barrero et al., 2021), while others '
        'document declines in collaborative output and innovation (Yang et al., 2022). These '
        'discrepancies may reflect differences in measurement approaches, industry contexts, '
        'and the maturity of remote work infrastructure within studied organizations.'
    )
    doc.add_paragraph(
        'Our research questions are: (1) How does work arrangement type (remote, hybrid, in-office) '
        'relate to self-reported and manager-assessed productivity? (2) What contextual factors '
        'moderate the relationship between remote work and productivity? (3) How do employees '
        'and team leads describe the mechanisms through which remote work affects daily output?'
    )

    # ---- Literature Review ----
    doc.add_heading('2. Literature Review', level=1)
    doc.add_paragraph(
        'The productivity implications of remote work have been studied through multiple theoretical '
        'lenses. Self-Determination Theory (Deci & Ryan, 1985) suggests that remote work may '
        'enhance intrinsic motivation by satisfying autonomy needs, while Social Exchange Theory '
        '(Blau, 1964) posits that flexible arrangements signal organizational trust, fostering '
        'reciprocal commitment. Job Demands-Resources Theory (Bakker & Demerouti, 2017) provides '
        'a framework for understanding how remote work simultaneously introduces new resources '
        '(schedule flexibility, reduced commute) and demands (isolation, technology dependence).'
    )
    doc.add_paragraph(
        'Meta-analytic evidence from Gajendran and Harrison (2007) covering 46 studies found a '
        'small but significant positive effect of telecommuting on productivity (d = 0.22). '
        'However, more recent meta-analyses incorporating pandemic-era data suggest the effect '
        'may be larger (d = 0.35) but highly heterogeneous across contexts (Wang et al., 2024). '
        'Industry-specific studies in technology sectors report higher productivity gains compared '
        'to service and manufacturing sectors (Atkinson & Sandiford, 2023).'
    )

    # ---- Methodology ----
    doc.add_heading('3. Methodology', level=1)

    doc.add_heading('3.1 Participants', level=2)
    doc.add_paragraph(
        'Participants were recruited from 14 mid-size technology companies (200-2,000 employees) '
        'headquartered in the United States and Canada. The final sample consisted of 1,247 '
        'knowledge workers: 412 fully remote, 498 hybrid (2-3 days in office), and 337 fully '
        'in-office. The sample was 54% female, 43% male, and 3% non-binary, with a mean age '
        'of 34.7 years (SD = 8.3). Ethnic composition was 58% White, 18% Asian, 12% Hispanic/Latino, '
        '8% Black, and 4% other/multiracial.'
    )

    doc.add_heading('3.2 Measures', level=2)
    doc.add_paragraph(
        'Productivity was assessed using three instruments: (a) the Workplace Productivity Scale '
        '(WPS; Koopmans et al., 2014), a 15-item self-report measure; (b) quarterly manager '
        'performance ratings on a 5-point scale; and (c) objective output metrics where available '
        '(e.g., code commits, tickets resolved, deliverables completed). Control variables included '
        'tenure, role seniority, team size, and prior remote work experience.'
    )

    doc.add_heading('3.3 Procedure', level=2)
    doc.add_paragraph(
        'Data collection occurred in four waves: baseline (January 2024), and follow-ups at '
        'months 4, 8, and 12. Surveys were administered via Qualtrics with a mean completion '
        'time of 18 minutes. Semi-structured interviews (45-60 minutes) were conducted via Zoom '
        'with 38 team leads purposively sampled to represent diversity in team size, function, '
        'and geographic distribution.'
    )

    # ---- Results ----
    doc.add_heading('4. Results', level=1)

    doc.add_heading('4.1 Quantitative Findings', level=2)
    doc.add_paragraph(
        'A one-way ANOVA revealed significant differences in WPS scores across work arrangement '
        'groups, F(2, 1244) = 18.67, p < .001, partial eta-squared = .029. Post-hoc Tukey HSD '
        'tests showed that fully remote workers (M = 4.12, SD = 0.87) scored significantly higher '
        'than hybrid workers (M = 3.78, SD = 0.93, p = .002) and in-office workers (M = 3.45, '
        'SD = 1.02, p < .001). The hybrid-to-in-office comparison was also significant (p = .014). '
        'Manager ratings showed a similar but smaller pattern, F(2, 1244) = 7.34, p < .001.'
    )

    doc.add_heading('4.2 Qualitative Findings', level=2)
    doc.add_paragraph(
        'Thematic analysis of interview transcripts identified three primary themes. First, '
        '"autonomy-driven focus" was the most frequently cited mechanism, with 29 of 38 leads '
        'noting that remote workers demonstrated longer uninterrupted work blocks. Second, '
        '"communication overhead" emerged as a counterbalancing factor: 22 leads reported that '
        'fully remote teams required more scheduled meetings, reducing net productive time. '
        'Third, "boundary management challenges" affected 15 leads\' teams, particularly those '
        'with caregiving responsibilities or shared living spaces.'
    )

    # ---- Discussion ----
    doc.add_heading('5. Discussion', level=1)
    doc.add_paragraph(
        'Our findings extend previous research by demonstrating sustained productivity advantages '
        'for remote workers over a 12-month period, even after controlling for self-selection bias '
        'through propensity score matching. The moderate effect size (partial eta-squared = .029) '
        'is consistent with recent meta-analytic estimates and suggests practically meaningful '
        'differences at the organizational level. The qualitative findings illuminate the mechanisms '
        'behind these gains, particularly the role of deep work periods enabled by reduced '
        'office interruptions (Newport, 2016).'
    )
    doc.add_paragraph(
        'Importantly, our results reveal significant moderating effects of organizational support. '
        'Companies providing structured communication protocols, dedicated home office stipends, '
        'and manager training in remote leadership saw 40% larger productivity gains among remote '
        'workers compared to companies without these supports. This underscores that remote work '
        'benefits are not automatic but depend on intentional organizational investment.'
    )

    # ---- Conclusion ----
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        'This study provides robust, longitudinal evidence that remote work arrangements are '
        'associated with higher self-reported and manager-assessed productivity among knowledge '
        'workers in the technology sector. However, these benefits are contingent on organizational '
        'support structures and individual self-regulation capacity. Organizations considering '
        'permanent remote or hybrid policies should invest in communication infrastructure, '
        'manager training, and employee wellness programs to maximize productivity outcomes.'
    )

    # ---- References ----
    doc.add_heading('References', level=1)
    references = [
        'Atkinson, C., & Sandiford, P. (2023). Remote work and productivity in technology firms: A sector analysis. Journal of Business Research, 156, 113-128.',
        'Bakker, A. B., & Demerouti, E. (2017). Job demands-resources theory: Taking stock and looking forward. Journal of Occupational Health Psychology, 22(3), 273-285.',
        'Barrero, J. M., Bloom, N., & Davis, S. J. (2021). Why working from home will stick. National Bureau of Economic Research, Working Paper 28731.',
        'Blau, P. M. (1964). Exchange and power in social life. Wiley.',
        'Bloom, N., Han, R., & Liang, J. (2023). How hybrid working from home works out. National Bureau of Economic Research, Working Paper 30292.',
        'Bloom, N., Liang, J., Roberts, J., & Ying, Z. J. (2015). Does working from home work? Evidence from a Chinese experiment. Quarterly Journal of Economics, 130(1), 165-218.',
        'Choudhury, P. (2022). Geographic mobility, immobility, and geographic flexibility: A review and agenda for research on the changing geography of work. Academy of Management Annals, 16(1), 258-302.',
        'Deci, E. L., & Ryan, R. M. (1985). Intrinsic motivation and self-determination in human behavior. Plenum.',
        'Gajendran, R. S., & Harrison, D. A. (2007). The good, the bad, and the unknown about telecommuting: Meta-analysis of psychological mediators and individual consequences. Journal of Applied Psychology, 92(6), 1524-1541.',
        'Koopmans, L., Bernaards, C. M., Hildebrandt, V. H., et al. (2014). Construct validity of the Individual Work Performance Questionnaire. Journal of Occupational and Environmental Medicine, 56(3), 331-337.',
        'Newport, C. (2016). Deep work: Rules for focused success in a distracted world. Grand Central Publishing.',
        'Wang, B., Liu, Y., Qian, J., & Parker, S. K. (2024). Achieving effective remote working during the COVID-19 pandemic: A work design perspective. Applied Psychology, 70(1), 16-59.',
        'Yang, L., Holtz, D., Jaffe, S., et al. (2022). The effects of remote work on collaboration among information workers. Nature Human Behaviour, 6(1), 43-54.',
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)  # hanging indent

    # ---- Appendix A (NO section break - this is the initial state) ----
    doc.add_heading('Appendix A: Survey Instrument', level=1)
    doc.add_paragraph(
        'The following items were used to assess workplace productivity. Participants rated '
        'each item on a 5-point Likert scale (1 = Strongly Disagree, 5 = Strongly Agree).'
    )

    items = [
        ('1.', 'I am able to complete my assigned tasks within expected timeframes.'),
        ('2.', 'The quality of my work output meets or exceeds expectations.'),
        ('3.', 'I effectively prioritize my daily work activities.'),
        ('4.', 'I maintain focus on important tasks without frequent distractions.'),
        ('5.', 'I contribute meaningfully to team projects and collaborative efforts.'),
        ('6.', 'I proactively identify and solve problems in my work.'),
        ('7.', 'I manage my workload effectively across competing deadlines.'),
        ('8.', 'I produce creative solutions to work challenges.'),
        ('9.', 'I communicate effectively with colleagues and supervisors.'),
        ('10.', 'I adapt quickly to changing work requirements and priorities.'),
        ('11.', 'I take initiative to improve work processes.'),
        ('12.', 'I deliver work that requires minimal revision.'),
        ('13.', 'I meet or exceed performance goals set by my manager.'),
        ('14.', 'I balance speed and accuracy in completing tasks.'),
        ('15.', 'I maintain consistent productivity levels throughout the work week.'),
    ]
    for num, item_text in items:
        p = doc.add_paragraph()
        run_num = p.add_run(num + ' ')
        run_num.bold = True
        p.add_run(item_text)

    doc.add_paragraph()  # spacing
    doc.add_paragraph(
        'Additional open-ended questions asked participants to describe: '
        '(a) their typical workday structure, (b) primary obstacles to productivity, and '
        '(c) strategies they use to maintain focus while working remotely.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
