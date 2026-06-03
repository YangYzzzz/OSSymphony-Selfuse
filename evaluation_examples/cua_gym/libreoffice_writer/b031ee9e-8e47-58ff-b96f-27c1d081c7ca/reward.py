"""
Reward Script: Insert section break before 'Appendix A' with different page style and 'Appendix' header
Task ID: writer_acad_041
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Document has >1 section (section break inserted before Appendix A)
  Component 2 (0.3): Appendix section header contains 'Appendix' text
  Component 3 (0.3): Appendix header is independent (not linked to previous) AND margins match body section
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_041'


def persist_app_state(domain):
    """Save any unsaved LibreOffice changes before verification."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for LibreOffice Writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Cannot import python-docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)
    print(f"INFO: Document has {num_sections} section(s)")

    # Component 1: Document has more than 1 section — meaning a section break was inserted (0.4 points)
    # In the initial document there is only 1 section. The task requires inserting a section break
    # before 'Appendix A', which creates a second section.
    try:
        if num_sections >= 2:
            # Further verify: the section break occurs before or at 'Appendix A' paragraph
            # Check that 'Appendix A' heading exists in the document after section break
            appendix_found = False
            for para in doc.paragraphs:
                if 'Appendix A' in para.text:
                    appendix_found = True
                    break

            if appendix_found:
                print(f"PASS: Component 1 -- Document has {num_sections} sections and 'Appendix A' exists (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 1 -- Document has {num_sections} sections but 'Appendix A' paragraph not found")
        else:
            print(f"FAIL: Component 1 -- Document has only {num_sections} section(s), expected >= 2")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: The appendix section's header contains 'Appendix' text (0.3 points)
    # In the initial document, section 0 has no meaningful header text.
    # The golden document has section 1 with header text 'Appendix'.
    try:
        if num_sections >= 2:
            # Check all sections beyond the first for a header containing 'Appendix'
            appendix_header_found = False
            for sec_idx in range(1, num_sections):
                section = doc.sections[sec_idx]
                header = section.header
                if header and header.paragraphs:
                    header_text = ' '.join(p.text for p in header.paragraphs).strip()
                    if 'appendix' in header_text.lower():
                        appendix_header_found = True
                        print(f"  INFO: Section {sec_idx} header text: {header_text!r}")
                        break

            if appendix_header_found:
                print(f"PASS: Component 2 -- Appendix section header contains 'Appendix' (0.3 pts)")
                total_score += 0.3
            else:
                # Gather debug info
                for sec_idx in range(1, num_sections):
                    section = doc.sections[sec_idx]
                    header = section.header
                    if header and header.paragraphs:
                        ht = ' '.join(p.text for p in header.paragraphs).strip()
                        print(f"  DEBUG: Section {sec_idx} header text: {ht!r}")
                print(f"FAIL: Component 2 -- No section header contains 'Appendix'")
        else:
            print(f"FAIL: Component 2 -- Only 1 section, cannot check appendix header")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Appendix header is independent (not linked to previous) AND margins match body (0.3 points)
    # The task says the new page style should use the same margins as the body.
    # The header must be independent (is_linked_to_previous=False) to display different text.
    try:
        if num_sections >= 2:
            body_section = doc.sections[0]
            body_margins = (
                body_section.left_margin,
                body_section.right_margin,
                body_section.top_margin,
                body_section.bottom_margin,
            )

            independent_and_matching = False
            for sec_idx in range(1, num_sections):
                section = doc.sections[sec_idx]
                header = section.header

                # Check header independence
                header_independent = not header.is_linked_to_previous
                print(f"  INFO: Section {sec_idx} header is_linked_to_previous: {header.is_linked_to_previous}")

                # Check margins match body
                sec_margins = (
                    section.left_margin,
                    section.right_margin,
                    section.top_margin,
                    section.bottom_margin,
                )
                margins_match = (body_margins == sec_margins)
                print(f"  INFO: Section {sec_idx} margins: L={section.left_margin} R={section.right_margin} T={section.top_margin} B={section.bottom_margin}")
                print(f"  INFO: Body margins:      L={body_section.left_margin} R={body_section.right_margin} T={body_section.top_margin} B={body_section.bottom_margin}")

                # Check that this section has 'Appendix' in its header
                header_text = ' '.join(p.text for p in header.paragraphs).strip() if header.paragraphs else ''
                if 'appendix' in header_text.lower():
                    if header_independent and margins_match:
                        independent_and_matching = True
                        break
                    elif not header_independent:
                        print(f"  DEBUG: Section {sec_idx} header is linked to previous")
                    elif not margins_match:
                        print(f"  DEBUG: Section {sec_idx} margins do not match body")

            if independent_and_matching:
                print(f"PASS: Component 3 -- Appendix header is independent and margins match body (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 3 -- Appendix header not independent or margins mismatch")
        else:
            print(f"FAIL: Component 3 -- Only 1 section, cannot check header independence/margins")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
