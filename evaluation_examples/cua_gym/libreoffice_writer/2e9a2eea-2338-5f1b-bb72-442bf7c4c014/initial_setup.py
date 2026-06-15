"""
Initial Setup: Set up a concordance file for alphabetical index of medical terms
Task ID: writer_mt_090
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_090'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# --- Medical content sections for a realistic reference document ---

CHAPTERS = [
    {
        "title": "Chapter 1: Cardiovascular System",
        "sections": [
            {
                "heading": "1.1 Hypertension",
                "content": [
                    "Hypertension, commonly known as high blood pressure, is a chronic medical condition in which the systemic arterial blood pressure is elevated. It is classified as either primary (essential) hypertension or secondary hypertension. Approximately 90-95% of cases are primary hypertension, defined as high blood pressure with no obvious underlying medical cause.",
                    "Persistent hypertension is one of the risk factors for strokes, heart attacks, heart failure, and arterial aneurysms. It is a leading cause of chronic kidney disease. Moderate elevation of arterial blood pressure leads to shortened life expectancy. Both dietary and lifestyle changes as well as antihypertensive medications can improve blood pressure control and decrease the risk of associated health complications.",
                    "The relationship between hypertension and atherosclerosis is well established. Patients with hypertension often present with elevated cholesterol levels and may develop coronary artery disease. Management of hypertension requires regular monitoring of blood pressure and adherence to prescribed medications including ACE inhibitors, beta-blockers, and diuretics.",
                ]
            },
            {
                "heading": "1.2 Arrhythmia and Cardiac Disorders",
                "content": [
                    "Arrhythmia refers to any irregularity in the heart's rhythm. Atrial fibrillation is the most common type of arrhythmia, affecting millions of patients worldwide. Bradycardia and tachycardia represent abnormally slow and fast heart rates respectively.",
                    "Myocardial infarction, commonly known as a heart attack, occurs when blood flow to a part of the heart muscle is blocked. This is typically caused by atherosclerosis in the coronary arteries. Symptoms include chest pain, shortness of breath, and diaphoresis. Emergency treatment with angioplasty or thrombolytic therapy is essential.",
                    "Heart failure is a chronic condition where the heart cannot pump blood efficiently. Patients may experience edema, fatigue, and dyspnea. Echocardiography is used to assess cardiac function, measuring the ejection fraction as a key indicator of heart performance.",
                ]
            },
            {
                "heading": "1.3 Vascular Conditions",
                "content": [
                    "Deep vein thrombosis (DVT) is the formation of a blood clot in a deep vein, usually in the legs. Pulmonary embolism can occur when a clot breaks free and travels to the lungs. Anticoagulant therapy with warfarin or heparin is the standard treatment.",
                    "Peripheral artery disease affects the blood vessels outside the heart and brain. Atherosclerosis is the most common cause, leading to reduced blood flow to the extremities. Patients may experience claudication and require revascularization procedures.",
                ]
            }
        ]
    },
    {
        "title": "Chapter 2: Endocrine System",
        "sections": [
            {
                "heading": "2.1 Diabetes Mellitus",
                "content": [
                    "Diabetes mellitus is a group of metabolic diseases characterized by chronic hyperglycemia resulting from defects in insulin secretion, insulin action, or both. Type 1 diabetes is caused by autoimmune destruction of the pancreatic beta cells, while Type 2 diabetes is characterized by insulin resistance and relative insulin deficiency.",
                    "Management of diabetes requires careful monitoring of blood glucose levels. Hemoglobin A1c (HbA1c) testing provides an average of blood sugar control over the past 2-3 months. Patients with diabetes are at increased risk for neuropathy, nephropathy, and retinopathy.",
                    "Diabetic ketoacidosis is a serious complication of diabetes that occurs when the body produces high levels of ketones. It is most common in Type 1 diabetes and requires immediate medical attention. Hypoglycemia, or low blood sugar, is another common complication that can result from excessive insulin administration.",
                ]
            },
            {
                "heading": "2.2 Thyroid Disorders",
                "content": [
                    "Hypothyroidism occurs when the thyroid gland does not produce enough thyroid hormone. Symptoms include fatigue, weight gain, cold intolerance, and depression. Hashimoto's thyroiditis is the most common cause of hypothyroidism in developed countries.",
                    "Hyperthyroidism is the overproduction of thyroid hormones. Graves' disease is the most common cause, presenting with symptoms such as weight loss, tremor, heat intolerance, and exophthalmos. Treatment options include antithyroid medications, radioactive iodine therapy, and thyroidectomy.",
                    "Thyroid nodules are common and usually benign. Fine-needle aspiration biopsy is used to evaluate suspicious nodules. Thyroid cancer, while relatively rare, requires surgical intervention followed by radioactive iodine ablation in most cases.",
                ]
            }
        ]
    },
    {
        "title": "Chapter 3: Respiratory System",
        "sections": [
            {
                "heading": "3.1 Asthma and COPD",
                "content": [
                    "Asthma is a chronic inflammatory disease of the airways characterized by variable and recurring symptoms, reversible airflow obstruction, and bronchospasm. Common symptoms include wheezing, coughing, chest tightness, and shortness of breath. Treatment includes inhaled corticosteroids and bronchodilators.",
                    "Chronic obstructive pulmonary disease (COPD) encompasses emphysema and chronic bronchitis. It is primarily caused by long-term exposure to irritating gases, most often from cigarette smoke. Spirometry testing is essential for diagnosis, measuring forced expiratory volume (FEV1) and forced vital capacity (FVC).",
                    "Pulmonary rehabilitation is an important component of COPD management. Patients benefit from exercise training, nutritional counseling, and education about their disease. Oxygen therapy may be required for patients with severe hypoxemia.",
                ]
            },
            {
                "heading": "3.2 Pneumonia and Infectious Diseases",
                "content": [
                    "Pneumonia is an infection that inflames the air sacs in one or both lungs. It can be caused by bacteria, viruses, or fungi. Streptococcus pneumoniae is the most common bacterial cause. Symptoms include cough with phlegm, fever, chills, and difficulty breathing.",
                    "Tuberculosis remains a significant global health concern. Mycobacterium tuberculosis infection primarily affects the lungs but can spread to other organs. Diagnosis involves chest X-ray, sputum culture, and the tuberculin skin test. Treatment requires a multi-drug regimen over 6-9 months.",
                    "Influenza is a contagious respiratory illness caused by influenza viruses. Antiviral medications such as oseltamivir can reduce the severity and duration of symptoms. Annual vaccination is recommended for high-risk populations.",
                ]
            }
        ]
    },
    {
        "title": "Chapter 4: Gastrointestinal System",
        "sections": [
            {
                "heading": "4.1 Gastroesophageal Reflux Disease",
                "content": [
                    "Gastroesophageal reflux disease (GERD) occurs when stomach acid frequently flows back into the esophagus. This acid reflux can irritate the lining of the esophagus, causing heartburn and regurgitation. Chronic GERD can lead to Barrett's esophagus, a precancerous condition.",
                    "Treatment of GERD includes lifestyle modifications, proton pump inhibitors (PPIs), and H2 receptor antagonists. Endoscopy may be required to evaluate the extent of esophageal damage. Fundoplication surgery is considered for patients who do not respond to medical therapy.",
                ]
            },
            {
                "heading": "4.2 Hepatic and Pancreatic Disorders",
                "content": [
                    "Cirrhosis is the result of long-term liver damage, leading to scarring (fibrosis) and liver failure. Common causes include chronic alcoholism, hepatitis B, and hepatitis C infection. Complications include portal hypertension, ascites, and hepatic encephalopathy.",
                    "Pancreatitis is inflammation of the pancreas. Acute pancreatitis is most commonly caused by gallstones or excessive alcohol consumption. Chronic pancreatitis can lead to malabsorption and diabetes. Serum amylase and lipase levels are elevated in acute pancreatitis.",
                    "Inflammatory bowel disease encompasses Crohn's disease and ulcerative colitis. These are chronic conditions characterized by inflammation of the gastrointestinal tract. Colonoscopy with biopsy is essential for diagnosis. Treatment includes aminosalicylates, corticosteroids, and immunomodulators.",
                ]
            }
        ]
    },
    {
        "title": "Chapter 5: Neurological System",
        "sections": [
            {
                "heading": "5.1 Cerebrovascular Disease",
                "content": [
                    "Stroke is a medical emergency that occurs when blood supply to part of the brain is interrupted (ischemic stroke) or when a blood vessel in the brain ruptures (hemorrhagic stroke). Rapid assessment using the NIH Stroke Scale and immediate CT imaging are critical for determining treatment.",
                    "Transient ischemic attack (TIA) is often called a mini-stroke and serves as a warning sign for future strokes. Carotid endarterectomy may be recommended for patients with significant carotid artery stenosis. Antiplatelet therapy with aspirin or clopidogrel is commonly prescribed for secondary prevention.",
                ]
            },
            {
                "heading": "5.2 Neurodegenerative Disorders",
                "content": [
                    "Alzheimer's disease is the most common cause of dementia, characterized by progressive cognitive decline and memory loss. Amyloid plaques and neurofibrillary tangles are the hallmark pathological features. Cholinesterase inhibitors may provide modest symptomatic benefit.",
                    "Parkinson's disease is a progressive movement disorder caused by degeneration of dopaminergic neurons in the substantia nigra. Cardinal symptoms include tremor, rigidity, bradykinesia, and postural instability. Levodopa remains the most effective medication for managing motor symptoms.",
                    "Multiple sclerosis is an autoimmune disease that affects the central nervous system. Demyelination of nerve fibers leads to various neurological symptoms including visual disturbances, muscle weakness, and coordination problems. MRI of the brain and spinal cord is essential for diagnosis.",
                    "Epilepsy is a neurological disorder characterized by recurrent seizures. Electroencephalography (EEG) is the primary diagnostic tool. Anticonvulsant medications such as valproate, carbamazepine, and levetiracetam are used for seizure control.",
                ]
            }
        ]
    },
    {
        "title": "Chapter 6: Musculoskeletal System",
        "sections": [
            {
                "heading": "6.1 Arthritis",
                "content": [
                    "Osteoarthritis is the most common form of arthritis, resulting from the breakdown of joint cartilage. It primarily affects weight-bearing joints such as the knees, hips, and spine. Treatment includes analgesics, physical therapy, and joint replacement surgery in severe cases.",
                    "Rheumatoid arthritis is an autoimmune disorder that causes chronic inflammation of the joints. Early diagnosis and treatment with disease-modifying antirheumatic drugs (DMARDs) such as methotrexate can slow disease progression. Biologic agents targeting TNF-alpha are used for refractory cases.",
                ]
            },
            {
                "heading": "6.2 Osteoporosis and Bone Health",
                "content": [
                    "Osteoporosis is characterized by decreased bone density and increased fracture risk. Dual-energy X-ray absorptiometry (DEXA) scanning is the standard method for measuring bone mineral density. Risk factors include aging, female sex, low body weight, and vitamin D deficiency.",
                    "Bisphosphonates such as alendronate are the first-line treatment for osteoporosis. Calcium and vitamin D supplementation are recommended for all patients at risk. Weight-bearing exercise and fall prevention strategies are important components of management.",
                ]
            }
        ]
    },
    {
        "title": "Chapter 7: Renal and Urological System",
        "sections": [
            {
                "heading": "7.1 Chronic Kidney Disease",
                "content": [
                    "Chronic kidney disease (CKD) is a progressive loss of renal function over months or years. The most common causes are diabetes and hypertension. Glomerular filtration rate (GFR) is used to classify the severity of CKD into five stages.",
                    "Management of CKD focuses on treating the underlying cause, controlling blood pressure, and managing complications such as anemia, metabolic acidosis, and hyperphosphatemia. Dialysis or kidney transplantation may be required for end-stage renal disease.",
                ]
            },
            {
                "heading": "7.2 Urinary Tract Conditions",
                "content": [
                    "Urinary tract infections (UTI) are among the most common infections. Escherichia coli is the most frequent causative organism. Symptoms include dysuria, frequency, and urgency. Antibiotics such as trimethoprim-sulfamethoxazole or nitrofurantoin are first-line treatments.",
                    "Nephrolithiasis (kidney stones) affects approximately 10% of the population. Calcium oxalate stones are the most common type. Diagnosis is typically made with CT scan. Treatment depends on stone size and location, ranging from conservative management to lithotripsy or surgical removal.",
                ]
            }
        ]
    },
    {
        "title": "Chapter 8: Oncology",
        "sections": [
            {
                "heading": "8.1 Common Malignancies",
                "content": [
                    "Lung cancer is the leading cause of cancer death worldwide. Non-small cell lung cancer accounts for approximately 85% of cases. Staging with PET-CT scanning guides treatment decisions. Treatment options include surgery, chemotherapy, radiation therapy, and targeted therapies such as EGFR inhibitors.",
                    "Breast cancer is the most common cancer in women. Mammography screening has improved early detection rates. Treatment is multimodal, including surgery (lumpectomy or mastectomy), radiation therapy, chemotherapy, and hormonal therapy for hormone receptor-positive tumors.",
                    "Colorectal cancer is the third most common cancer globally. Colonoscopy screening with polypectomy can prevent cancer development. The TNM staging system guides prognosis and treatment planning. Adjuvant chemotherapy with FOLFOX regimen is standard for stage III disease.",
                ]
            },
            {
                "heading": "8.2 Hematologic Malignancies",
                "content": [
                    "Leukemia encompasses a group of cancers arising from the bone marrow. Acute lymphoblastic leukemia (ALL) is the most common childhood cancer, while acute myeloid leukemia (AML) is more common in adults. Flow cytometry and cytogenetic analysis are essential for classification.",
                    "Lymphoma is classified as Hodgkin's or non-Hodgkin's lymphoma. Excisional lymph node biopsy is required for diagnosis. PET-CT scanning is used for staging. Treatment protocols vary but may include chemotherapy (ABVD for Hodgkin's, R-CHOP for diffuse large B-cell lymphoma), radiation, and stem cell transplantation.",
                ]
            }
        ]
    },
    {
        "title": "Chapter 9: Infectious Diseases",
        "sections": [
            {
                "heading": "9.1 Bacterial Infections",
                "content": [
                    "Sepsis is a life-threatening condition caused by the body's response to infection. Early recognition and treatment with broad-spectrum antibiotics are critical. The Sequential Organ Failure Assessment (SOFA) score is used to identify patients at risk of poor outcomes.",
                    "Methicillin-resistant Staphylococcus aureus (MRSA) is a significant cause of hospital-acquired infections. Vancomycin remains the drug of choice for serious MRSA infections. Contact precautions and hand hygiene are essential for preventing transmission.",
                    "Clostridium difficile infection is increasingly common, particularly in hospitalized patients receiving antibiotics. Symptoms range from mild diarrhea to severe colitis. Treatment includes oral vancomycin or fidaxomicin. Fecal microbiota transplantation is effective for recurrent infections.",
                ]
            },
            {
                "heading": "9.2 Viral Infections",
                "content": [
                    "Human immunodeficiency virus (HIV) infection leads to acquired immunodeficiency syndrome (AIDS) if untreated. Antiretroviral therapy (ART) has transformed HIV from a fatal disease to a manageable chronic condition. CD4 count and viral load monitoring guide treatment decisions.",
                    "Hepatitis B and hepatitis C are major causes of chronic liver disease worldwide. Hepatitis B vaccination is highly effective for prevention. Direct-acting antiviral agents can cure hepatitis C in over 95% of cases. Screening for hepatitis is recommended for high-risk populations.",
                ]
            }
        ]
    },
    {
        "title": "Chapter 10: Dermatology and Allergic Conditions",
        "sections": [
            {
                "heading": "10.1 Common Skin Conditions",
                "content": [
                    "Psoriasis is a chronic autoimmune condition that causes rapid buildup of skin cells, resulting in scaling on the skin surface. Topical corticosteroids are the first-line treatment for mild to moderate psoriasis. Phototherapy and systemic agents including biologics are used for moderate to severe cases.",
                    "Eczema (atopic dermatitis) is characterized by dry, itchy, and inflamed skin. It commonly begins in childhood and may persist into adulthood. Management includes emollients, topical corticosteroids, and avoidance of triggers. Immunosuppressants may be needed for severe cases.",
                    "Melanoma is the most dangerous form of skin cancer. The ABCDE criteria (asymmetry, border, color, diameter, evolution) aid in early detection. Excisional biopsy with adequate margins is the primary treatment. Immunotherapy with checkpoint inhibitors has improved outcomes for advanced melanoma.",
                ]
            },
            {
                "heading": "10.2 Allergic and Immunologic Disorders",
                "content": [
                    "Anaphylaxis is a severe, potentially life-threatening allergic reaction. Epinephrine (adrenaline) is the first-line treatment and should be administered immediately. Common triggers include foods, insect stings, and medications. Patients should carry auto-injectable epinephrine at all times.",
                    "Systemic lupus erythematosus (SLE) is a chronic autoimmune disease that can affect multiple organ systems. The hallmark butterfly rash across the cheeks is characteristic. Diagnosis is based on clinical criteria and laboratory findings including antinuclear antibody (ANA) testing. Treatment includes hydroxychloroquine, corticosteroids, and immunosuppressants.",
                ]
            }
        ]
    },
]


def create_initial():
    doc = Document()

    # Title page
    title = doc.add_heading('Medical Reference Handbook', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Comprehensive Guide to Clinical Medicine')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    edition = doc.add_paragraph()
    edition.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = edition.add_run('Fourth Edition, 2025')
    run.font.size = Pt(12)
    run.italic = True

    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = authors.add_run('Editors: Dr. Elena Vasquez, MD, PhD\nDr. James Harrington, MD, FACP\nDr. Priya Sharma, MD, MPH')
    run.font.size = Pt(11)

    doc.add_page_break()

    # Table of Contents placeholder
    toc_heading = doc.add_heading('Table of Contents', level=1)
    for ch in CHAPTERS:
        p = doc.add_paragraph(ch['title'], style='List Number')
        for sec in ch['sections']:
            doc.add_paragraph(sec['heading'], style='List Bullet')
    doc.add_page_break()

    # Preface
    doc.add_heading('Preface', level=1)
    doc.add_paragraph(
        'This Medical Reference Handbook is designed as a comprehensive clinical '
        'resource for medical students, residents, and practicing physicians. It covers '
        'major organ systems and disease processes with emphasis on pathophysiology, '
        'diagnosis, and evidence-based management. The content has been reviewed by '
        'specialists in cardiology, endocrinology, pulmonology, gastroenterology, '
        'neurology, oncology, and infectious diseases.'
    )
    doc.add_paragraph(
        'Each chapter provides detailed discussions of common and important conditions, '
        'including their clinical presentation, diagnostic workup, and treatment options. '
        'Cross-references between related conditions such as hypertension and chronic '
        'kidney disease, or diabetes and cardiovascular complications, help readers '
        'understand the interconnected nature of human physiology and disease.'
    )
    doc.add_page_break()

    # Build chapters
    for ch in CHAPTERS:
        doc.add_heading(ch['title'], level=1)
        for sec in ch['sections']:
            doc.add_heading(sec['heading'], level=2)
            for para_text in sec['content']:
                doc.add_paragraph(para_text)
        doc.add_page_break()

    # Appendix with medical abbreviations table
    doc.add_heading('Appendix A: Common Medical Abbreviations', level=1)
    abbreviations = [
        ('ABG', 'Arterial Blood Gas'),
        ('ACE', 'Angiotensin-Converting Enzyme'),
        ('ANA', 'Antinuclear Antibody'),
        ('BMI', 'Body Mass Index'),
        ('BMP', 'Basic Metabolic Panel'),
        ('CBC', 'Complete Blood Count'),
        ('CKD', 'Chronic Kidney Disease'),
        ('COPD', 'Chronic Obstructive Pulmonary Disease'),
        ('CT', 'Computed Tomography'),
        ('DVT', 'Deep Vein Thrombosis'),
        ('ECG', 'Electrocardiogram'),
        ('EEG', 'Electroencephalography'),
        ('GERD', 'Gastroesophageal Reflux Disease'),
        ('GFR', 'Glomerular Filtration Rate'),
        ('HbA1c', 'Hemoglobin A1c'),
        ('HIV', 'Human Immunodeficiency Virus'),
        ('INR', 'International Normalized Ratio'),
        ('MRI', 'Magnetic Resonance Imaging'),
        ('NSAID', 'Non-Steroidal Anti-Inflammatory Drug'),
        ('PET', 'Positron Emission Tomography'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Abbreviation'
    hdr[1].text = 'Full Term'
    for run in hdr[0].paragraphs[0].runs:
        run.bold = True
    for run in hdr[1].paragraphs[0].runs:
        run.bold = True
    for abbr, full in abbreviations:
        row = table.add_row().cells
        row[0].text = abbr
        row[1].text = full

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
