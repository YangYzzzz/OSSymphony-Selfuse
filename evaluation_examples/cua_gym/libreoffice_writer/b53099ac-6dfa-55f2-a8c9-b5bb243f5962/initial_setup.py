"""
Initial Setup: Create a business letter document with sender and recipient addresses.
Task ID: writer_rd_050
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_050'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set standard letter page size (8.5 x 11)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Sender address block (top of letter)
    sender_lines = [
        "TechVision Inc.",
        "456 Oak Avenue",
        "Chicago, IL 60601",
    ]
    for line in sender_lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = "Liberation Sans"
        run.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)

    # Date
    doc.add_paragraph()  # blank line
    date_para = doc.add_paragraph()
    run = date_para.add_run("March 15, 2026")
    run.font.name = "Liberation Sans"
    run.font.size = Pt(11)

    # Blank line
    doc.add_paragraph()

    # Recipient address block
    recipient_lines = [
        "John Smith",
        "Acme Corp",
        "123 Main Street",
        "Springfield, IL 62701",
    ]
    for line in recipient_lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = "Liberation Sans"
        run.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)

    # Blank line
    doc.add_paragraph()

    # Salutation
    salutation = doc.add_paragraph()
    run = salutation.add_run("Dear Mr. Smith,")
    run.font.name = "Liberation Sans"
    run.font.size = Pt(11)

    # Blank line
    doc.add_paragraph()

    # Body paragraphs
    body_text_1 = (
        "Thank you for your continued partnership with TechVision Inc. We are pleased to "
        "inform you that the Q1 2026 product roadmap has been finalized, and we would like "
        "to schedule a meeting to discuss the upcoming integration milestones for the Acme "
        "Corp deployment."
    )
    para1 = doc.add_paragraph()
    run = para1.add_run(body_text_1)
    run.font.name = "Liberation Sans"
    run.font.size = Pt(11)

    body_text_2 = (
        "As outlined in our previous correspondence, the Phase 2 rollout is expected to "
        "begin on April 15, 2026. Our engineering team has completed the preliminary testing "
        "and we are confident that the new features will meet the performance benchmarks "
        "agreed upon in the service level agreement."
    )
    para2 = doc.add_paragraph()
    run = para2.add_run(body_text_2)
    run.font.name = "Liberation Sans"
    run.font.size = Pt(11)

    body_text_3 = (
        "Please let us know your availability for a call next week. We suggest either "
        "Tuesday at 2:00 PM or Thursday at 10:00 AM Central Time. Our project manager, "
        "Elena Rodriguez, will be joining the call along with the lead developer, David Park."
    )
    para3 = doc.add_paragraph()
    run = para3.add_run(body_text_3)
    run.font.name = "Liberation Sans"
    run.font.size = Pt(11)

    # Blank line
    doc.add_paragraph()

    # Closing
    closing = doc.add_paragraph()
    run = closing.add_run("Sincerely,")
    run.font.name = "Liberation Sans"
    run.font.size = Pt(11)

    # Blank lines for signature space
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block
    sig_lines = [
        "Michael Chen",
        "Director of Client Solutions",
        "TechVision Inc.",
        "mchen@techvision.com",
        "(312) 555-0147",
    ]
    for line in sig_lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = "Liberation Sans"
        run.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
