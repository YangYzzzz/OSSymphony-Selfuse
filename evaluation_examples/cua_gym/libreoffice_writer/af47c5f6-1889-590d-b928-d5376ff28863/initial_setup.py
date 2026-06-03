"""
Initial Setup: Insert firm logo into letter header
Task ID: writer_legal_082
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_082'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
LOGO_PATH = f'{WORKDIR}/firm_logo.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_logo():
    """Create a simple firm logo PNG image."""
    img = Image.new('RGBA', (200, 100), (255, 255, 255, 0))
    from PIL import ImageDraw, ImageFont
    draw = ImageDraw.Draw(img)
    # Draw a blue rectangle as a simple logo mark
    draw.rectangle([10, 10, 90, 90], fill=(25, 60, 120, 255))
    # Add "M&A" text inside
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 28)
    except (IOError, OSError):
        font = ImageFont.load_default()
    draw.text((20, 30), "M&A", fill=(255, 255, 255, 255), font=font)
    # Add a gold accent bar
    draw.rectangle([100, 40, 190, 55], fill=(194, 155, 61, 255))
    img.save(LOGO_PATH)
    print(f'Logo created: {LOGO_PATH}')


def create_initial():
    create_logo()

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.25)
    section.bottom_margin = Inches(1.0)

    # --- Header: centered firm name only (NO logo) ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    run = hp.add_run("Mitchell & Associates, LLP")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(25, 60, 120)
    hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Footer ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fr = fp.add_run("1200 Commerce Tower, Suite 450 | Chicago, IL 60601 | (312) 555-0198")
    fr.font.name = "Times New Roman"
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(128, 128, 128)

    # --- Body: Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(6)
    dr = date_para.add_run("March 28, 2026")
    dr.font.name = "Times New Roman"
    dr.font.size = Pt(11)

    # --- Recipient address ---
    recipient_lines = [
        "Ms. Catherine Reynolds",
        "Vice President of Operations",
        "Lakewood Manufacturing Corp.",
        "4500 Industrial Parkway",
        "Milwaukee, WI 53201",
    ]
    for line in recipient_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

    # Blank line
    doc.add_paragraph()

    # --- Salutation ---
    sal = doc.add_paragraph()
    sr = sal.add_run("Dear Ms. Reynolds,")
    sr.font.name = "Times New Roman"
    sr.font.size = Pt(11)

    # --- Body paragraphs ---
    body_texts = [
        "Thank you for retaining Mitchell & Associates, LLP regarding the pending "
        "breach of contract dispute with Hartfield Supply Co. This letter summarizes "
        "our initial assessment and recommended course of action following our "
        "consultation on March 25, 2026.",

        "After reviewing the executed supply agreement dated September 12, 2024, "
        "the relevant invoices, and your correspondence with Hartfield's procurement "
        "department, it is our opinion that Lakewood Manufacturing has strong grounds "
        "for a breach of contract claim under Wisconsin Commercial Code Section 402.711. "
        "The evidence indicates that Hartfield failed to deliver conforming goods on "
        "three separate occasions between November 2025 and February 2026, resulting "
        "in documented damages exceeding $287,000.",

        "We recommend initiating the dispute resolution process by sending a formal "
        "demand letter to Hartfield Supply Co. within the next fourteen business days. "
        "Should they fail to respond or offer an inadequate remedy, we are prepared to "
        "file a civil complaint in Milwaukee County Circuit Court. Our preliminary "
        "damages calculation, inclusive of consequential losses from production delays, "
        "amounts to approximately $342,500.",

        "Please note that the statute of limitations for this type of claim under "
        "Wisconsin law is six years from the date of breach. However, prompt action "
        "will preserve your position and demonstrate diligence to the court.",

        "Our fee arrangement, as discussed, will be structured on an hourly basis at "
        "the rates outlined in the enclosed engagement letter. We estimate total legal "
        "costs for the pre-litigation phase at $15,000 to $22,000, depending on "
        "Hartfield's responsiveness.",

        "We look forward to representing Lakewood Manufacturing in this matter. Please "
        "do not hesitate to contact me directly at (312) 555-0198 or via email at "
        "j.mitchell@mitchellassociates.com should you have any questions.",
    ]

    for text in body_texts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

    # --- Closing ---
    doc.add_paragraph()
    closing = doc.add_paragraph()
    cr = closing.add_run("Sincerely,")
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(11)

    # Signature space
    doc.add_paragraph()
    doc.add_paragraph()

    sig_name = doc.add_paragraph()
    snr = sig_name.add_run("James T. Mitchell, Esq.")
    snr.font.name = "Times New Roman"
    snr.font.size = Pt(11)
    snr.bold = True

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_after = Pt(0)
    str_ = sig_title.add_run("Managing Partner")
    str_.font.name = "Times New Roman"
    str_.font.size = Pt(11)

    sig_firm = doc.add_paragraph()
    sig_firm.paragraph_format.space_before = Pt(0)
    sfr = sig_firm.add_run("Mitchell & Associates, LLP")
    sfr.font.name = "Times New Roman"
    sfr.font.size = Pt(11)

    # --- Enclosure note ---
    doc.add_paragraph()
    enc = doc.add_paragraph()
    er = enc.add_run("Encl: Engagement Letter")
    er.font.name = "Times New Roman"
    er.font.size = Pt(10)
    er.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
