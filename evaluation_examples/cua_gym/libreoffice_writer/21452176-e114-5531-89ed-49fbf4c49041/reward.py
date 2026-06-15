"""
Reward Script: Technical Specification Document for API Gateway Service v3.0
Task ID: writer_wf_051
Domain: libreoffice_writer
Scoring:
  Component 1: Title present with correct text (0.10)
  Component 2: 6 Heading 1 sections with required names (0.20)
  Component 3: Document Control table structure (0.15)
  Component 4: System Requirements table (4 component rows, 3 cols) (0.15)
  Component 5: API Endpoints table (6 endpoint rows, 4 cols) (0.15)
  Component 6: Error Codes table (5 error rows, 2 cols) (0.15)
  Component 7: TOC section exists + body content is non-trivial (0.10)
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_051'

# Required Heading 1 section names (case-insensitive matching)
REQUIRED_H1_SECTIONS = [
    'overview',
    'system requirements',
    'api endpoints',
    'error codes',
    'security',
    'deployment notes',
]


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather paragraphs and tables upfront
    paragraphs = doc.paragraphs
    tables = doc.tables

    # --- Component 1: Title contains "API Gateway Service v3.0" (0.10 pts) ---
    try:
        title_paras = [p for p in paragraphs if p.style and p.style.name == 'Title']
        title_texts = [p.text.strip().lower() for p in title_paras]
        if any('api gateway service' in t and 'v3' in t for t in title_texts):
            print(f"PASS: Component 1 — Title found: {title_paras[0].text!r} (0.10 pts)")
            total_score += 0.10
        else:
            # Also check Heading 0 or first paragraph as fallback
            first_text = paragraphs[0].text.strip().lower() if paragraphs else ''
            if 'api gateway service' in first_text and 'v3' in first_text:
                print(f"PASS: Component 1 — Title found in first paragraph (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 1 — No title containing 'API Gateway Service v3.0'. Found title paras: {title_texts}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # --- Component 2: 6 Heading 1 sections with required names (0.20 pts) ---
    try:
        h1_paras = [p.text.strip().lower() for p in paragraphs if p.style and p.style.name == 'Heading 1']
        matched = 0
        for req in REQUIRED_H1_SECTIONS:
            if any(req in h for h in h1_paras):
                matched += 1
            else:
                print(f"  MISS: Heading 1 '{req}' not found")
        if matched == 6:
            print(f"PASS: Component 2 — All 6 Heading 1 sections found (0.20 pts)")
            total_score += 0.20
        elif matched >= 4:
            partial = round(0.20 * (matched / 6), 2)
            print(f"PARTIAL: Component 2 — {matched}/6 Heading 1 sections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {matched}/6 Heading 1 sections. Found: {h1_paras}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # --- Component 3: Document Control table (0.15 pts) ---
    # Must have 5 columns: Version, Date, Author, Reviewer, Status
    try:
        dc_table = None
        required_dc_headers = {'version', 'date', 'author', 'reviewer', 'status'}
        for t in tables:
            if len(t.columns) == 5:
                headers = {c.text.strip().lower() for c in t.rows[0].cells}
                if required_dc_headers.issubset(headers):
                    dc_table = t
                    break
        if dc_table is not None:
            # Must have at least 1 data row (header + 1)
            if len(dc_table.rows) >= 2:
                print(f"PASS: Component 3 — Document Control table with {len(dc_table.rows)-1} data rows, 5 cols (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 3 — Document Control table found but no data rows")
        else:
            print(f"FAIL: Component 3 — No table with headers {required_dc_headers}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # --- Component 4: System Requirements table (4 component rows, 3 cols) (0.15 pts) ---
    try:
        sr_table = None
        required_sr_headers = {'component', 'minimum', 'recommended'}
        for t in tables:
            if len(t.columns) == 3:
                headers = {c.text.strip().lower() for c in t.rows[0].cells}
                if required_sr_headers.issubset(headers):
                    sr_table = t
                    break
        if sr_table is not None:
            data_rows = len(sr_table.rows) - 1  # minus header
            if data_rows >= 4:
                print(f"PASS: Component 4 — System Requirements table with {data_rows} component rows (0.15 pts)")
                total_score += 0.15
            elif data_rows >= 2:
                partial = round(0.15 * (data_rows / 4), 2)
                print(f"PARTIAL: Component 4 — System Requirements table has {data_rows}/4 rows ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — System Requirements table has only {data_rows} data rows")
        else:
            print(f"FAIL: Component 4 — No System Requirements table (3 cols: Component, Minimum, Recommended)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # --- Component 5: API Endpoints table (6 endpoint rows, 4 cols) (0.15 pts) ---
    try:
        ep_table = None
        required_ep_headers = {'method', 'endpoint', 'description'}
        for t in tables:
            if len(t.columns) == 4:
                headers = {c.text.strip().lower() for c in t.rows[0].cells}
                if required_ep_headers.issubset(headers):
                    ep_table = t
                    break
        if ep_table is not None:
            data_rows = len(ep_table.rows) - 1
            if data_rows >= 6:
                print(f"PASS: Component 5 — API Endpoints table with {data_rows} endpoint rows (0.15 pts)")
                total_score += 0.15
            elif data_rows >= 3:
                partial = round(0.15 * (data_rows / 6), 2)
                print(f"PARTIAL: Component 5 — API Endpoints table has {data_rows}/6 rows ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 — API Endpoints table has only {data_rows} data rows")
        else:
            print(f"FAIL: Component 5 — No API Endpoints table (4 cols: Method, Endpoint, Description, Auth Required)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # --- Component 6: Error Codes table (5 error rows, 2 cols) (0.15 pts) ---
    try:
        ec_table = None
        for t in tables:
            if len(t.columns) == 2:
                headers = {c.text.strip().lower() for c in t.rows[0].cells}
                if 'code' in headers and 'message' in headers:
                    ec_table = t
                    break
        if ec_table is not None:
            data_rows = len(ec_table.rows) - 1
            if data_rows >= 5:
                print(f"PASS: Component 6 — Error Codes table with {data_rows} error rows (0.15 pts)")
                total_score += 0.15
            elif data_rows >= 3:
                partial = round(0.15 * (data_rows / 5), 2)
                print(f"PARTIAL: Component 6 — Error Codes table has {data_rows}/5 rows ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 6 — Error Codes table has only {data_rows} data rows")
        else:
            print(f"FAIL: Component 6 — No Error Codes table (2 cols: Code, Message)")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # --- Component 7: TOC section + non-trivial body content (0.10 pts) ---
    try:
        # Check for TOC heading (Heading 1 or Heading 2 containing "Table of Contents" or "TOC")
        toc_found = False
        for p in paragraphs:
            if p.style and 'Heading' in p.style.name:
                if 'table of contents' in p.text.strip().lower() or 'toc' == p.text.strip().lower():
                    toc_found = True
                    break

        # Check for non-trivial body: at least 5 normal paragraphs with text
        normal_paras_with_text = [p for p in paragraphs if p.style and p.style.name == 'Normal' and len(p.text.strip()) > 20]
        has_body = len(normal_paras_with_text) >= 3

        if toc_found and has_body:
            print(f"PASS: Component 7 — TOC section found + {len(normal_paras_with_text)} body paragraphs (0.10 pts)")
            total_score += 0.10
        elif toc_found:
            print(f"PARTIAL: Component 7 — TOC found but insufficient body content ({len(normal_paras_with_text)} paras) (0.05 pts)")
            total_score += 0.05
        elif has_body:
            print(f"PARTIAL: Component 7 — Body content present but no TOC heading (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 7 — No TOC section and insufficient body content")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
