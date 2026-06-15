"""
Initial Setup: Internal audit report for page break + heading style task
Task ID: writer_para_067
Domain: libreoffice_writer

Creates audit_report.docx with 11 paragraphs.
Paragraphs 5 and 9 are in Default Paragraph Style with NO page_break_before.
The agent must:
  - Apply 'Heading 1' style to paragraphs 5 and 9
  - Set page_break_before=True on paragraphs 5 and 9
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_para_067'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Title — already Heading 1 style
    p1 = doc.add_paragraph('Internal Audit Report 2024', style='Heading 1')
    p1.paragraph_format.page_break_before = False

    # Paragraph 2: Introduction line 1
    p2 = doc.add_paragraph(
        'This report presents the findings of the annual internal audit '
        'conducted from October to December 2024.'
    )
    p2.paragraph_format.page_break_before = False

    # Paragraph 3: Scope description
    p3 = doc.add_paragraph(
        'The audit scope included financial controls, IT security, and '
        'regulatory compliance across all four regional offices.'
    )
    p3.paragraph_format.page_break_before = False

    # Paragraph 4: Summary findings
    p4 = doc.add_paragraph(
        'We identified 23 findings, of which 5 are classified as high priority.'
    )
    p4.paragraph_format.page_break_before = False

    # Paragraph 5: Section heading — MUST be Default Paragraph Style, NO page break
    p5 = doc.add_paragraph('Financial Controls Assessment')
    p5.paragraph_format.page_break_before = False
    # Explicitly set to Normal / Default paragraph style (already default)

    # Paragraph 6: Financial controls overview
    p6 = doc.add_paragraph(
        'The financial controls framework was found to be generally effective. '
        'However, three significant gaps were identified in the accounts payable process.'
    )
    p6.paragraph_format.page_break_before = False

    # Paragraph 7: Finding F-01
    p7 = doc.add_paragraph(
        'Finding F-01: Duplicate payment detection controls were bypassed for '
        'transactions under $1,000.'
    )
    p7.paragraph_format.page_break_before = False

    # Paragraph 8: Finding F-02
    p8 = doc.add_paragraph(
        'Finding F-02: Authorization thresholds for expense reimbursements '
        'were not consistently enforced.'
    )
    p8.paragraph_format.page_break_before = False

    # Paragraph 9: Section heading — MUST be Default Paragraph Style, NO page break
    p9 = doc.add_paragraph('IT Security Assessment')
    p9.paragraph_format.page_break_before = False
    # Explicitly set to Normal / Default paragraph style (already default)

    # Paragraph 10: IT security overview
    p10 = doc.add_paragraph(
        'The IT security posture has improved significantly since the last audit. '
        'Penetration testing revealed two medium-severity vulnerabilities.'
    )
    p10.paragraph_format.page_break_before = False

    # Paragraph 11: Finding IT-01
    p11 = doc.add_paragraph(
        'Finding IT-01: The staging environment database contained unmasked production data.'
    )
    p11.paragraph_format.page_break_before = False

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer with DISPLAY=:0
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
