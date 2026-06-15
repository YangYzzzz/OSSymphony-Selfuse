"""
Initial Setup: Insert arbitration clause text frame on signature page
Task ID: writer_legal_060
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_060'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard letter size with 1" margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title ---
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Effective Date: March 15, 2025')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacer

    # --- Parties ---
    parties_heading = doc.add_heading('RECITALS', level=1)

    parties_text = (
        'This Professional Services Agreement ("Agreement") is entered into as of '
        'March 15, 2025, by and between Meridian Technology Solutions, Inc., a Delaware '
        'corporation with its principal place of business at 4200 Harbor Boulevard, Suite 300, '
        'Newport Beach, California 92660 ("Company"), and Catalyst Engineering Group, LLC, '
        'a California limited liability company with its principal place of business at '
        '1875 Century Park East, Suite 700, Los Angeles, California 90067 ("Contractor").'
    )
    p = doc.add_paragraph(parties_text)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    whereas_items = [
        'WHEREAS, Company desires to engage Contractor to provide certain professional '
        'engineering and technical consulting services as described herein; and',
        'WHEREAS, Contractor represents that it has the necessary expertise, qualifications, '
        'and resources to perform such services; and',
        'WHEREAS, the parties wish to establish the terms and conditions under which '
        'Contractor will provide such services to Company.',
    ]
    for item in whereas_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    p = doc.add_paragraph(
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements contained '
        'herein, and for other good and valuable consideration, the receipt and sufficiency '
        'of which are hereby acknowledged, the parties agree as follows:'
    )
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # --- Section 1: Scope of Services ---
    doc.add_heading('1. SCOPE OF SERVICES', level=1)

    scope_paras = [
        '1.1 Contractor shall provide professional engineering consultation, technical '
        'analysis, and project management services as outlined in Exhibit A attached hereto '
        'and incorporated by reference ("Services").',
        '1.2 Contractor shall assign qualified personnel with appropriate expertise to '
        'perform the Services. Key personnel include Dr. Rebecca Torres (Lead Engineer), '
        'James Whitfield (Senior Analyst), and Patricia Nakamura (Project Coordinator).',
        '1.3 The Services shall be performed at Company\'s facilities located at 4200 Harbor '
        'Boulevard, Suite 300, Newport Beach, California 92660, unless otherwise agreed in '
        'writing by the parties.',
        '1.4 Contractor shall devote such time and attention to the Services as is reasonably '
        'necessary for the satisfactory completion of the work within the agreed-upon timeline.',
    ]
    for text in scope_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # --- Section 2: Compensation ---
    doc.add_heading('2. COMPENSATION AND PAYMENT', level=1)

    comp_paras = [
        '2.1 Company shall pay Contractor a total fee of Two Hundred Forty-Five Thousand '
        'Dollars ($245,000.00) for the complete performance of the Services, payable in '
        'monthly installments as follows:',
        '    (a) An initial payment of $49,000.00 upon execution of this Agreement;',
        '    (b) Four subsequent monthly payments of $49,000.00 each, due on the first '
        'business day of each calendar month following the commencement date.',
        '2.2 All invoices shall be submitted to Company\'s Accounts Payable department at '
        'ap@meridiantech.com and shall include a detailed description of Services performed '
        'during the billing period.',
        '2.3 Payment shall be due within thirty (30) days of receipt of a proper invoice. '
        'Late payments shall accrue interest at the rate of 1.5% per month or the maximum '
        'rate permitted by applicable law, whichever is less.',
    ]
    for text in comp_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # --- Section 3: Term and Termination ---
    doc.add_heading('3. TERM AND TERMINATION', level=1)

    term_paras = [
        '3.1 This Agreement shall commence on March 15, 2025, and shall continue for a '
        'period of six (6) months, unless earlier terminated in accordance with this Section 3.',
        '3.2 Either party may terminate this Agreement for convenience upon thirty (30) days\' '
        'prior written notice to the other party.',
        '3.3 Either party may terminate this Agreement immediately upon written notice if the '
        'other party: (a) materially breaches any provision of this Agreement and fails to '
        'cure such breach within fifteen (15) days after receipt of written notice thereof; or '
        '(b) becomes insolvent, files for bankruptcy protection, or makes an assignment for '
        'the benefit of creditors.',
        '3.4 Upon termination, Contractor shall deliver to Company all work product, '
        'documentation, and materials developed in connection with the Services.',
    ]
    for text in term_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # --- Section 4: Confidentiality ---
    doc.add_heading('4. CONFIDENTIALITY', level=1)

    conf_paras = [
        '4.1 Each party acknowledges that in the course of performing under this Agreement, '
        'it may receive or have access to confidential and proprietary information of the '
        'other party ("Confidential Information").',
        '4.2 Each party agrees to hold all Confidential Information in strict confidence and '
        'not to disclose such information to any third party without the prior written consent '
        'of the disclosing party, except as required by law.',
        '4.3 The obligations of confidentiality shall survive the termination of this Agreement '
        'for a period of three (3) years.',
    ]
    for text in conf_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # --- Section 5: Intellectual Property ---
    doc.add_heading('5. INTELLECTUAL PROPERTY', level=1)

    ip_paras = [
        '5.1 All work product, deliverables, inventions, and materials created by Contractor '
        'in the performance of the Services shall be the sole and exclusive property of Company.',
        '5.2 Contractor hereby assigns to Company all rights, title, and interest in and to '
        'all intellectual property created in connection with the Services.',
    ]
    for text in ip_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # --- Section 6: Governing Law ---
    doc.add_heading('6. GOVERNING LAW', level=1)

    p = doc.add_paragraph(
        '6.1 This Agreement shall be governed by and construed in accordance with the laws '
        'of the State of California, without regard to its conflict of laws provisions.'
    )
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # --- Section 7: Miscellaneous ---
    doc.add_heading('7. MISCELLANEOUS', level=1)

    misc_paras = [
        '7.1 This Agreement constitutes the entire agreement between the parties with respect '
        'to the subject matter hereof and supersedes all prior and contemporaneous agreements '
        'and understandings, whether written or oral.',
        '7.2 This Agreement may not be amended or modified except by a written instrument '
        'signed by both parties.',
        '7.3 If any provision of this Agreement is held to be invalid or unenforceable, the '
        'remaining provisions shall continue in full force and effect.',
        '7.4 This Agreement may be executed in counterparts, each of which shall be deemed an '
        'original, but all of which together shall constitute one and the same instrument.',
    ]
    for text in misc_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # --- Signature Page ---
    # Add page break before signature page
    doc.add_page_break()

    sig_heading = doc.add_heading('SIGNATURE PAGE', level=1)
    sig_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph(
        'IN WITNESS WHEREOF, the parties have executed this Professional Services Agreement '
        'as of the date first written above.'
    )
    p.paragraph_format.space_after = Pt(24)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # Company signature block
    company_lines = [
        'MERIDIAN TECHNOLOGY SOLUTIONS, INC.',
        '',
        '',
        '________________________________________',
        'Name: Jonathan R. Blackwell',
        'Title: Chief Executive Officer',
        'Date: _____________________',
    ]
    for line in company_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            if line == 'MERIDIAN TECHNOLOGY SOLUTIONS, INC.':
                run.bold = True

    doc.add_paragraph()  # spacer

    # Contractor signature block
    contractor_lines = [
        'CATALYST ENGINEERING GROUP, LLC',
        '',
        '',
        '________________________________________',
        'Name: Dr. Angela M. Reeves',
        'Title: Managing Director',
        'Date: _____________________',
    ]
    for line in contractor_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            if line == 'CATALYST ENGINEERING GROUP, LLC':
                run.bold = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
