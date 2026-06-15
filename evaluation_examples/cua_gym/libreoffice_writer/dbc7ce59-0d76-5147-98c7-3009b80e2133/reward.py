"""
FINAL REWARD SCRIPT - SUCCESS
Task: Working on a physics lab report in LibreOffice Writer. After dropping in a formula, I need to tag it with a caption that says exactly "Equation 1" (Category: Equation, Position: Below the object) so it lines up with the rest of my numbered equations later. Can you walk me through the clicks to make that happen?
Generated: 2025-09-10 17:47:25
Status: success
Model: azure-o3
Total Steps: 14
"""

import os
import re
from docx import Document

"""
Reward Script for Physics-Lab-Report Task
----------------------------------------
Task to verify:
1. A caption paragraph that reads exactly "Equation 1" must exist.
2. That caption must appear *below* the formula it describes, i.e. the previous non-empty
   paragraph should look like a mathematical expression (simple heuristic: it contains
   one of the characters =, <, >, ^ ).

Scoring (progressive):
• 0.6 points  – Caption "Equation 1" found as its own paragraph.
• 0.4 points  – The immediately preceding non-empty paragraph appears to be a formula,
                meaning the caption was inserted below the object it describes.
Total possible = 1.0.

The script purposefully gives *no* points for merely loading the file or other
natural conditions.  Each awarded point is tied to a concrete verification step.
"""


# ---------- Helper: locate the document to inspect ----------

def find_document_path() -> str | None:
    """Return absolute path of the .docx the agent worked on (if any)."""
    # Known target name from task description
    preferred = (
        "/home/user/working_on_a_physics_lab_report_in_libreoffice_writer_after_dropping_in_a_formula_i_need_to_tag_it_w.docx",
    )
    for p in preferred:
        if os.path.exists(p):
            return p

    # Fallback: first .docx in /home/user
    for fname in os.listdir("/home/user"):
        if fname.lower().endswith(".docx"):
            return os.path.join("/home/user", fname)
    return None


# ---------- Core verification logic ----------

def verify_equation_caption(doc: Document) -> float:
    """Return a partial score (0-1 range) based on caption presence and position."""
    caption_index: int | None = None

    # 1) Look for caption paragraph with exact text "Equation 1"
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == "equation 1":
            caption_index = i
            break

    if caption_index is None:
        print("✗ Caption 'Equation 1' not found as its own paragraph")
        return 0.0  # Nothing else to check if caption missing

    score = 0.6  # Caption found earns 0.6
    print(f"✓ Found caption paragraph at index {caption_index} (0.6 pts)")

    # 2) Ensure immediately previous non-empty paragraph resembles a formula
    prev_text: str | None = None
    for j in range(caption_index - 1, -1, -1):
        text = doc.paragraphs[j].text.strip()
        if text:  # first non-empty above the caption
            prev_text = text
            break

    if prev_text is None:
        print("✗ No non-empty paragraph before the caption – cannot confirm positioning")
        return score  # only caption points awarded

    # Heuristic: formula likely contains one of these characters
    if re.search(r'[=<>&^+\-*/]', prev_text):
        print(f"✓ Previous paragraph appears to be a formula: {prev_text!r} (0.4 pts)")
        score += 0.4
    else:
        print(f"✗ Previous paragraph does not look like a formula: {prev_text!r}")

    return score


def verify_task() -> float:
    """Main entry – orchestrates the verification and prints the reward."""
    file_path = find_document_path()
    if not file_path:
        print("✗ No DOCX file found for verification")
        print("REWARD: 0.0")
        return 0.0

    print(f"Analyzing file: {file_path}")
    try:
        doc = Document(file_path)
    except Exception as exc:
        print(f"✗ Failed to load document: {exc}")
        print("REWARD: 0.0")
        return 0.0

    # Perform the actual checks and accumulate score
    score = verify_equation_caption(doc)
    final_score = min(score, 1.0)  # Defensive cap

    print(f"REWARD: {final_score}")
    return final_score


# Execute when run as a script
if __name__ == "__main__":
    verify_task()

