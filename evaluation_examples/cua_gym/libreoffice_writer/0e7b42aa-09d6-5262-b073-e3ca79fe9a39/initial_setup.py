#!/usr/bin/env python3
"""initial_setup.py - Creates the initial Writer document for writer_legal_025.

The document is a realistic contract ending with the closing paragraph:
"IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above."
No signature block exists.
"""

import subprocess
import shlex
import os
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

OUTPUT_PATH = "/home/user/writer_legal_025.docx"


def create_document():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run("SERVICES AGREEMENT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    # --- Date line ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.space_after = Pt(18)
    run = date_para.add_run("Effective Date: January 15, 2026")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    # Helper to add a body paragraph
    def add_body(text, space_after=Pt(6), bold_label=None):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = space_after
        if bold_label:
            r = p.add_run(bold_label)
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = "Times New Roman"
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
        return p

    # --- Preamble ---
    add_body(
        'This Services Agreement (the "Agreement") is entered into as of the Effective Date '
        'set forth above, by and between ABC Corporation, a Delaware corporation with its '
        'principal offices located at 123 Main Street, Suite 400, Wilmington, DE 19801 '
        '("ABC"), and XYZ Industries, a California corporation with its principal offices '
        'located at 456 Innovation Drive, San Jose, CA 95112 ("XYZ").',
        space_after=Pt(12),
    )

    # --- Recitals ---
    add_body("RECITALS", space_after=Pt(6), bold_label="")
    add_body(
        "WHEREAS, ABC is engaged in the business of providing technology consulting and "
        "software development services; and"
    )
    add_body(
        "WHEREAS, XYZ desires to engage ABC to perform certain professional services as "
        "described herein; and"
    )
    add_body(
        "WHEREAS, the parties desire to set forth the terms and conditions under which ABC "
        "will provide such services to XYZ;",
        space_after=Pt(12),
    )
    add_body(
        "NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth "
        "herein, and for other good and valuable consideration, the receipt and sufficiency "
        "of which are hereby acknowledged, the parties agree as follows:",
        space_after=Pt(12),
    )

    # --- Section 1 ---
    add_body(
        " Scope of Services. ABC shall provide to XYZ the professional services described "
        'in Exhibit A attached hereto (the "Services"). ABC shall perform the Services in '
        "a professional and workmanlike manner, consistent with industry standards.",
        bold_label="1.",
    )

    # --- Section 2 ---
    add_body(
        " Term. This Agreement shall commence on the Effective Date and shall continue for "
        "a period of twelve (12) months, unless earlier terminated in accordance with "
        "Section 7 of this Agreement.",
        bold_label="2.",
    )

    # --- Section 3 ---
    add_body(
        " Compensation. In consideration for the Services, XYZ shall pay ABC the fees set "
        "forth in Exhibit B. Payment shall be due within thirty (30) days of receipt of "
        "each invoice. Late payments shall bear interest at a rate of one and one-half "
        "percent (1.5%) per month.",
        bold_label="3.",
    )

    # --- Section 4 ---
    add_body(
        " Confidentiality. Each party agrees to hold in confidence all Confidential "
        "Information received from the other party. Confidential Information shall not be "
        "disclosed to any third party without the prior written consent of the disclosing "
        "party, except as required by law.",
        bold_label="4.",
    )

    # --- Section 5 ---
    add_body(
        " Intellectual Property. All intellectual property rights in any work product "
        "created by ABC in the performance of the Services shall be owned by XYZ upon full "
        "payment of all fees due hereunder.",
        bold_label="5.",
    )

    # --- Section 6 ---
    add_body(
        " Limitation of Liability. In no event shall either party be liable to the other "
        "for any indirect, incidental, special, consequential, or punitive damages arising "
        "out of or related to this Agreement, regardless of the form of action.",
        bold_label="6.",
    )

    # --- Section 7 ---
    add_body(
        " Termination. Either party may terminate this Agreement upon sixty (60) days' "
        "prior written notice to the other party. In the event of a material breach, the "
        "non-breaching party may terminate immediately upon written notice.",
        bold_label="7.",
    )

    # --- Section 8 ---
    add_body(
        " Governing Law. This Agreement shall be governed by and construed in accordance "
        "with the laws of the State of Delaware, without regard to its conflict of laws "
        "provisions.",
        bold_label="8.",
        space_after=Pt(12),
    )

    # --- Closing paragraph ---
    closing = doc.add_paragraph()
    closing.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    closing.paragraph_format.space_before = Pt(12)
    closing.paragraph_format.space_after = Pt(0)
    run = closing.add_run(
        "IN WITNESS WHEREOF, the parties have executed this Agreement as of the date "
        "first written above."
    )
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    doc.save(OUTPUT_PATH)
    print(f"Document saved to {OUTPUT_PATH}")


def launch_gui(command, delay_sec=2.0):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


if __name__ == "__main__":
    create_document()
    launch_gui(f'libreoffice --writer "{OUTPUT_PATH}"', delay_sec=2.0)
    print("Done. Document opened in LibreOffice Writer.")
