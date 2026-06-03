"""
Initial Setup: Create a report document with a TOC that uses dot leaders
Task ID: writer_mt_080
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_080'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc_field(doc):
    """Add a TOC field code that LibreOffice can update."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar_separate)

    # Add static TOC entries with dot leaders (these are the visible entries before update)
    toc_entries = [
        (1, "Executive Summary", "2"),
        (1, "Market Analysis", "3"),
        (2, "Industry Overview", "3"),
        (2, "Competitive Landscape", "5"),
        (2, "Target Demographics", "7"),
        (1, "Financial Projections", "9"),
        (2, "Revenue Forecast", "9"),
        (2, "Cost Structure", "11"),
        (3, "Fixed Costs", "11"),
        (3, "Variable Costs", "12"),
        (1, "Implementation Strategy", "13"),
        (2, "Phase 1: Foundation", "13"),
        (2, "Phase 2: Growth", "15"),
        (2, "Phase 3: Optimization", "17"),
        (1, "Risk Assessment", "19"),
        (1, "Conclusions and Recommendations", "21"),
    ]

    for level, title, page in toc_entries:
        toc_para = doc.add_paragraph()
        toc_para.style = doc.styles[f'TOC Heading'] if level == 0 else None

        # Set indentation based on level
        indent = Inches(0.25 * (level - 1))
        toc_para.paragraph_format.left_indent = indent

        # Add right-aligned tab stop with DOT leader
        tab_stops = toc_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        # Entry text + tab + page number
        toc_para.add_run(title)
        toc_para.add_run("\t")
        toc_para.add_run(page)

    # End field
    end_para = doc.add_paragraph()
    run_end = end_para.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end._element.append(fldChar_end)

    return paragraph


def create_initial():
    doc = Document()

    # -- Title Page --
    title = doc.add_heading("Minimalist Report", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Quarterly Business Review — Q1 2025")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run("Prepared by: Elena Vasquez, Senior Strategy Analyst")
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_page_break()

    # -- Table of Contents --
    toc_title = doc.add_heading("Table of Contents", level=1)
    add_toc_field(doc)

    doc.add_page_break()

    # -- Section 1: Executive Summary --
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report presents a comprehensive analysis of our market position "
        "and strategic direction for the upcoming fiscal year. Key findings indicate "
        "a 12.3% growth opportunity in the mid-market segment, driven by increasing "
        "demand for automated workflow solutions."
    )
    doc.add_paragraph(
        "Revenue for Q1 2025 reached $4.87 million, exceeding projections by 8.2%. "
        "Customer acquisition costs decreased by 15% quarter-over-quarter, while "
        "the average contract value increased to $23,400 per annum."
    )

    doc.add_page_break()

    # -- Section 2: Market Analysis --
    doc.add_heading("Market Analysis", level=1)

    doc.add_heading("Industry Overview", level=2)
    doc.add_paragraph(
        "The enterprise software market continues its robust expansion, with total "
        "addressable market estimated at $672 billion globally. Cloud-native solutions "
        "now represent 58% of new deployments, up from 41% in 2023."
    )

    doc.add_heading("Competitive Landscape", level=2)
    doc.add_paragraph(
        "Our primary competitors — TechVantage Corp, NovaSoft Industries, and "
        "Meridian Solutions — have collectively invested $340 million in R&D during "
        "the past fiscal year. However, our Net Promoter Score of 72 remains the "
        "highest in the segment."
    )

    doc.add_heading("Target Demographics", level=2)
    doc.add_paragraph(
        "Analysis of our customer base reveals three primary segments: mid-market "
        "enterprises (250-1000 employees) accounting for 54% of revenue, small "
        "businesses (50-249 employees) at 28%, and large enterprises (1000+) at 18%."
    )

    doc.add_page_break()

    # -- Section 3: Financial Projections --
    doc.add_heading("Financial Projections", level=1)

    doc.add_heading("Revenue Forecast", level=2)
    doc.add_paragraph(
        "Based on current growth trajectories and pipeline analysis, we project "
        "annual revenue of $21.4 million for FY2025, representing a 23% increase "
        "over the previous fiscal year."
    )

    doc.add_heading("Cost Structure", level=2)

    doc.add_heading("Fixed Costs", level=3)
    doc.add_paragraph(
        "Fixed operational costs, including facilities, core staffing, and "
        "infrastructure, are projected at $8.2 million annually. This represents "
        "a 5% increase due to planned office expansion in Austin, Texas."
    )

    doc.add_heading("Variable Costs", level=3)
    doc.add_paragraph(
        "Variable costs — encompassing cloud infrastructure scaling, contractor "
        "fees, and performance-based compensation — are estimated at $4.1 million, "
        "scaling proportionally with revenue growth."
    )

    doc.add_page_break()

    # -- Section 4: Implementation Strategy --
    doc.add_heading("Implementation Strategy", level=1)

    doc.add_heading("Phase 1: Foundation", level=2)
    doc.add_paragraph(
        "Q1-Q2 2025: Establish core platform enhancements, complete API v3.0 "
        "migration, and onboard 15 strategic partner integrations. Budget allocation: "
        "$2.4 million."
    )

    doc.add_heading("Phase 2: Growth", level=2)
    doc.add_paragraph(
        "Q3 2025: Launch targeted marketing campaigns in EMEA and APAC regions. "
        "Expand sales team by 8 representatives focused on the mid-market segment. "
        "Expected customer acquisition: 120 new accounts."
    )

    doc.add_heading("Phase 3: Optimization", level=2)
    doc.add_paragraph(
        "Q4 2025: Implement advanced analytics dashboard, introduce AI-powered "
        "customer success recommendations, and optimize pricing tiers based on "
        "usage pattern analysis from the first three quarters."
    )

    doc.add_page_break()

    # -- Section 5: Risk Assessment --
    doc.add_heading("Risk Assessment", level=1)
    doc.add_paragraph(
        "Key risk factors include potential regulatory changes in data privacy "
        "legislation (particularly in the EU), supply chain disruptions affecting "
        "hardware procurement timelines, and talent acquisition challenges in "
        "the machine learning engineering domain."
    )
    doc.add_paragraph(
        "Mitigation strategies include maintaining a 90-day compliance buffer "
        "for regulatory changes, diversifying our cloud provider portfolio across "
        "three major vendors, and establishing university partnership programs "
        "for pipeline development."
    )

    doc.add_page_break()

    # -- Section 6: Conclusions --
    doc.add_heading("Conclusions and Recommendations", level=1)
    doc.add_paragraph(
        "The organization is well-positioned for sustained growth in FY2025. "
        "We recommend prioritizing the mid-market expansion strategy, investing "
        "in product-led growth capabilities, and accelerating the partner "
        "ecosystem development. With disciplined execution, achieving the $21.4M "
        "revenue target is within reach."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
