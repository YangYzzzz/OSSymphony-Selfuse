"""
Reward Script: Create a research paper in LibreOffice Writer
Task ID: writer_wf_007
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15): Title present with correct text
  Component 2 (0.25): All 8 section headings present as Heading 1
  Component 3 (0.15): Abstract paragraph is italicized
  Component 4 (0.15): Results table with >= 5 data rows
  Component 5 (0.10): References section has >= 3 entries
  Component 6 (0.10): Double-spaced line spacing
  Component 7 (0.10): Times New Roman 12pt font
"""

import os

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_007'

EXPECTED_SECTIONS = [
    'Abstract',
    'Introduction',
    'Literature Review',
    'Methodology',
    'Results',
    'Discussion',
    'Conclusion',
    'References',
]


def persist_app_state():
    """Try to save any unsaved LibreOffice document via Ctrl+S."""
    try:
        os.environ["DISPLAY"] = ":0"
        import pyautogui
        import time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraphs and their styles for reuse
    all_paras = doc.paragraphs
    heading1_texts = []
    for p in all_paras:
        if p.style and p.style.name == 'Heading 1':
            heading1_texts.append(p.text.strip())

    # =========================================================
    # Component 1: Title present with correct text (0.15 points)
    # =========================================================
    try:
        title_found = False
        for p in all_paras:
            if p.style and p.style.name in ('Title', 'Heading 0'):
                if 'impact of remote work' in p.text.lower() and 'productivity' in p.text.lower():
                    title_found = True
                    break
        # Also accept Heading 1 as first heading if it matches the title
        if not title_found and all_paras:
            first_text = all_paras[0].text.strip().lower()
            if 'impact of remote work' in first_text and 'productivity' in first_text:
                title_found = True

        if title_found:
            print(f"PASS: Component 1 — Title found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Title not found or incorrect")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # =========================================================
    # Component 2: All 8 section headings present as Heading 1 (0.25 points)
    # =========================================================
    try:
        found_sections = []
        for expected in EXPECTED_SECTIONS:
            for h in heading1_texts:
                if expected.lower() == h.lower():
                    found_sections.append(expected)
                    break

        section_ratio = len(found_sections) / len(EXPECTED_SECTIONS)
        # Award proportional credit: all 8 = 0.25, partial = proportional
        if section_ratio == 1.0:
            print(f"PASS: Component 2 — All 8 sections found as Heading 1: {found_sections} (0.25 pts)")
            total_score += 0.25
        elif section_ratio > 0:
            partial = round(0.25 * section_ratio, 3)
            missing = [s for s in EXPECTED_SECTIONS if s not in found_sections]
            print(f"PARTIAL: Component 2 — {len(found_sections)}/8 sections found. Missing: {missing} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No Heading 1 sections found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =========================================================
    # Component 3: Abstract paragraph is italicized (0.15 points)
    # =========================================================
    try:
        # Find the Abstract heading index, then check the paragraph(s) after it
        abstract_idx = None
        for i, p in enumerate(all_paras):
            if p.style and p.style.name == 'Heading 1' and p.text.strip().lower() == 'abstract':
                abstract_idx = i
                break

        abstract_italic = False
        if abstract_idx is not None:
            # Check paragraphs between Abstract heading and next heading
            for j in range(abstract_idx + 1, len(all_paras)):
                p = all_paras[j]
                if p.style and 'Heading' in p.style.name:
                    break  # reached next section
                if p.text.strip():
                    # Check if all runs are italic
                    runs_with_text = [r for r in p.runs if r.text.strip()]
                    if runs_with_text and all(r.font.italic for r in runs_with_text):
                        abstract_italic = True
                    break  # only check first non-empty paragraph

        if abstract_italic:
            print(f"PASS: Component 3 — Abstract paragraph is italicized (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 — Abstract paragraph is not italicized")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # =========================================================
    # Component 4: Results table with >= 5 data rows (0.15 points)
    # =========================================================
    try:
        tables = doc.tables
        table_found = False
        if len(tables) > 0:
            # Find a table with at least 6 rows (1 header + 5 data)
            for table in tables:
                num_rows = len(table.rows)
                if num_rows >= 6:
                    table_found = True
                    print(f"PASS: Component 4 — Table found with {num_rows} rows (>= 6 = header + 5 data) (0.15 pts)")
                    total_score += 0.15
                    break
            if not table_found:
                # Accept table with at least 5 rows (could be 5 data rows without header, or header + 4)
                for table in tables:
                    num_rows = len(table.rows)
                    if num_rows >= 5:
                        table_found = True
                        partial = 0.10
                        print(f"PARTIAL: Component 4 — Table found with {num_rows} rows (partial credit) ({partial} pts)")
                        total_score += partial
                        break
        if not table_found:
            if len(tables) > 0:
                print(f"FAIL: Component 4 — Table(s) found but none with >= 5 data rows. Largest: {max(len(t.rows) for t in tables)} rows")
            else:
                print(f"FAIL: Component 4 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # =========================================================
    # Component 5: References section has >= 3 entries (0.10 points)
    # =========================================================
    try:
        # Find References heading and count non-empty paragraphs after it
        refs_idx = None
        for i, p in enumerate(all_paras):
            if p.style and p.style.name == 'Heading 1' and p.text.strip().lower() == 'references':
                refs_idx = i
                break

        ref_count = 0
        if refs_idx is not None:
            for j in range(refs_idx + 1, len(all_paras)):
                p = all_paras[j]
                if p.style and 'Heading' in p.style.name:
                    break  # next section
                if p.text.strip():
                    ref_count += 1

        if ref_count >= 3:
            print(f"PASS: Component 5 — {ref_count} references found (>= 3) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — Only {ref_count} reference(s) found, expected >= 3")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # =========================================================
    # Component 6: Double-spaced line spacing (0.10 points)
    # =========================================================
    try:
        # Check that body paragraphs (Normal style) have line_spacing ~2.0
        normal_paras = [p for p in all_paras if p.style and p.style.name == 'Normal' and p.text.strip()]
        if normal_paras:
            double_spaced_count = 0
            for p in normal_paras:
                ls = p.paragraph_format.line_spacing
                if ls is not None and abs(float(ls) - 2.0) < 0.1:
                    double_spaced_count += 1

            ratio = double_spaced_count / len(normal_paras) if normal_paras else 0
            if ratio >= 0.8:
                print(f"PASS: Component 6 — {double_spaced_count}/{len(normal_paras)} paragraphs double-spaced (0.10 pts)")
                total_score += 0.10
            elif ratio > 0:
                partial = round(0.10 * ratio, 3)
                print(f"PARTIAL: Component 6 — {double_spaced_count}/{len(normal_paras)} paragraphs double-spaced ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 6 — No paragraphs have double spacing")
        else:
            print(f"FAIL: Component 6 — No Normal-style paragraphs found to check spacing")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # =========================================================
    # Component 7: Times New Roman 12pt font (0.10 points)
    # =========================================================
    try:
        # Check runs in Normal paragraphs for Times New Roman 12pt
        checked_runs = 0
        matching_runs = 0
        for p in all_paras:
            if p.style and p.style.name == 'Normal' and p.text.strip():
                for r in p.runs:
                    if not r.text.strip():
                        continue
                    checked_runs += 1
                    font_ok = r.font.name and 'times new roman' in r.font.name.lower()
                    size_ok = r.font.size is not None and abs(r.font.size.pt - 12.0) < 0.5
                    if font_ok and size_ok:
                        matching_runs += 1

        if checked_runs > 0:
            ratio = matching_runs / checked_runs
            if ratio >= 0.8:
                print(f"PASS: Component 7 — {matching_runs}/{checked_runs} runs are Times New Roman 12pt (0.10 pts)")
                total_score += 0.10
            elif ratio > 0:
                partial = round(0.10 * ratio, 3)
                print(f"PARTIAL: Component 7 — {matching_runs}/{checked_runs} runs match ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 7 — No runs have Times New Roman 12pt font")
        else:
            print(f"FAIL: Component 7 — No runs found to check font")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
