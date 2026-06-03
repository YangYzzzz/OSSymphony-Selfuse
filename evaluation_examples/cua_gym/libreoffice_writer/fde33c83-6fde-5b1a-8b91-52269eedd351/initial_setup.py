"""
Initial Setup: Create a Writer document with title and subtitle for style creation task.
Task ID: writer_bs_092
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_092'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading("Urban Transportation Networks", level=1)

    # --- Subtitle (Default Paragraph Style - NOT styled as Subtitle) ---
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle_para.add_run("A Comprehensive Analysis of Modern Approaches")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Body content for realism ---
    doc.add_paragraph()

    intro = doc.add_heading("1. Introduction", level=2)

    p1 = doc.add_paragraph(
        "The rapid expansion of urban areas worldwide has placed tremendous pressure "
        "on transportation infrastructure. Cities across the globe are grappling with "
        "challenges ranging from traffic congestion and air pollution to equitable "
        "access and sustainable mobility. This report examines the evolving landscape "
        "of urban transportation, analyzing both established systems and emerging "
        "innovations that promise to reshape how people move through metropolitan areas."
    )

    p2 = doc.add_paragraph(
        "According to the United Nations, approximately 68% of the world's population "
        "will live in urban areas by 2050. This demographic shift demands a fundamental "
        "rethinking of transportation planning, with particular emphasis on multimodal "
        "integration, environmental sustainability, and digital transformation."
    )

    doc.add_heading("2. Current Challenges", level=2)

    # Add a table of challenges
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"

    headers = ["Challenge", "Impact Level", "Primary Affected Areas"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    data = [
        ["Traffic Congestion", "Critical", "Downtown corridors, highway interchanges"],
        ["Air Quality Degradation", "High", "Dense urban cores, industrial zones"],
        ["Infrastructure Aging", "High", "Bridge networks, rail systems"],
        ["Equity in Access", "Medium-High", "Suburban fringes, low-income neighborhoods"],
        ["Funding Shortfalls", "Medium", "Capital projects, maintenance budgets"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()

    doc.add_heading("3. Emerging Solutions", level=2)

    solutions = [
        "Electric and autonomous vehicle fleets integrated with public transit",
        "Mobility-as-a-Service (MaaS) platforms consolidating ticketing and routing",
        "Dedicated cycling infrastructure and micro-mobility sharing programs",
        "Smart traffic management using real-time sensor data and AI optimization",
        "Transit-oriented development policies linking housing with rail stations",
    ]
    for sol in solutions:
        doc.add_paragraph(sol, style="List Bullet")

    doc.add_paragraph()

    doc.add_heading("4. Case Studies", level=2)

    p3 = doc.add_paragraph(
        "Singapore's Land Transport Authority has implemented a comprehensive "
        "Electronic Road Pricing system that dynamically adjusts tolls based on "
        "real-time congestion levels. Since its deployment in 2018, peak-hour "
        "traffic volumes have decreased by approximately 15%, while public transit "
        "ridership has grown by 8% annually."
    )

    p4 = doc.add_paragraph(
        "Copenhagen's ambitious cycling strategy has resulted in over 450 kilometers "
        "of dedicated bicycle lanes, with 62% of residents commuting by bicycle as of "
        "2024. The city's investment of approximately EUR 150 million in cycling "
        "infrastructure over the past decade has yielded estimated annual savings of "
        "EUR 230 million in healthcare and productivity gains."
    )

    doc.add_heading("5. Conclusions", level=2)

    p5 = doc.add_paragraph(
        "The future of urban transportation lies at the intersection of technological "
        "innovation, policy reform, and community engagement. Cities that successfully "
        "integrate these elements will be better positioned to provide efficient, "
        "equitable, and environmentally sustainable mobility options for their growing "
        "populations."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
