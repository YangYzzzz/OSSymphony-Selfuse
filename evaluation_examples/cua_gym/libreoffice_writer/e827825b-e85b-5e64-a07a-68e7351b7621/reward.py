"""
FINAL REWARD SCRIPT - SUCCESS
Task: Table 1 in my Writer document wound up with a fifth column that’s completely blank. How do I get rid of just that last (5th) column without disturbing the rest of the table?
Generated: 2025-09-10 17:25:48
Status: success
Model: azure-o3
Total Steps: 2
"""

from docx import Document
import os


def verify_table_column_removed(file_path: str) -> float:
    """Verify that the 5th (blank) column was removed from Table 1 in the document.

    Scoring (progressive – max 1.0):
    1. 0.7 pts – Table 1 now contains **4 or fewer** columns (so the former 5th column is gone).
    2. 0.3 pts – After removal, **no column is completely blank** (ensures the only-blank column was really removed).
    Returns a float in the range [0.0, 1.0].
    """

    print(f"Verifying document: {file_path}")
    score = 0.0
    MAX_SCORE = 1.0

    # ---------- 1.  Basic file/table checks (no points awarded) ----------
    if not os.path.exists(file_path):
        print("✗ File not found – verification failed")
        return 0.0  # cannot continue

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Unable to open DOCX: {e}")
        return 0.0

    if not doc.tables:
        print("✗ No tables found – expected at least one table")
        return 0.0

    table = doc.tables[0]  # Table 1 as per task statement

    # ---------- 2.  Requirement 1 – 5th column removed (0.7 pts) ----------
    row_col_counts = [len(row.cells) for row in table.rows]
    max_cols = max(row_col_counts)
    min_cols = min(row_col_counts)
    print("Row-wise column counts:", row_col_counts)

    if max_cols <= 4:  # all rows must be ≤ 4 columns now
        print("✓ Table has 4 or fewer columns – 5th column successfully removed (0.7 pts)")
        score += 0.7

        # ---------- 3.  Requirement 2 – no column completely blank (0.3 pts) ----------
        current_cols = max_cols  # they should all match, but we take the max safely
        blank_columns = []
        for col_idx in range(current_cols):
            if all(
                (len(row.cells) > col_idx and not row.cells[col_idx].text.strip())
                for row in table.rows
            ):
                blank_columns.append(col_idx)

        if blank_columns:
            print(f"✗ Found entirely blank column(s) at 0-based index/indices: {blank_columns}")
        else:
            print("✓ No completely blank columns remain (0.3 pts)")
            score += 0.3
    else:
        print("✗ Table still contains 5 or more columns – 5th column not removed")

    final_score = min(score, MAX_SCORE)
    print(f"REWARD: {final_score}")
    return final_score


# ----------------- Execute verification when script is run -----------------
if __name__ == "__main__":
    FILE_PATH = "/home/user/table_1_in_my_writer_document_wound_up_with_a_fifth_column_thats_completely_blank_how_do_i_get_rid_o.docx"
    verify_table_column_removed(FILE_PATH)
