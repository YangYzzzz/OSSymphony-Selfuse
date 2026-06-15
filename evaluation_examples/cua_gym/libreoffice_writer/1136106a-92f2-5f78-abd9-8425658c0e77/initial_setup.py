"""
Initial Setup: Create a document with straight quotes for curly-quote conversion task
Task ID: writer_tech_073
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_073'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Quarterly Software Development Report", level=0)

    # --- Section 1: Executive Summary ---
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        'This report provides an overview of the "Horizon" project\'s progress '
        'during Q1 2025. The development team has achieved "significant milestones" '
        'in both the front-end and back-end components. As the lead architect noted, '
        '"We\'ve exceeded our sprint velocity targets by 15% this quarter."'
    )
    doc.add_paragraph(
        'Key highlights include the launch of the "Smart Dashboard" feature, '
        'the migration to a "microservices" architecture, and the completion '
        'of the "Phase 2" security audit. The team\'s motto remains: '
        '"Ship quality code, ship it often."'
    )

    # --- Section 2: Technical Achievements ---
    doc.add_heading("Technical Achievements", level=1)

    doc.add_heading("API Gateway Redesign", level=2)
    doc.add_paragraph(
        'The API gateway, previously described as a "monolithic bottleneck," '
        'has been redesigned using the "strangler fig" pattern. Sarah Chen\'s '
        'team completed the transition in just three sprints. The new gateway '
        'supports "rate limiting," "circuit breaking," and "request routing" '
        'with sub-millisecond overhead.'
    )
    doc.add_paragraph(
        'According to the performance benchmarks, the new system handles '
        '"10,000 requests per second" compared to the previous "2,500 requests '
        'per second." Marcus Johnson commented: "It\'s a night-and-day difference."'
    )

    doc.add_heading("Database Optimization", level=2)
    doc.add_paragraph(
        'The database team identified several "query anti-patterns" that were '
        'causing performance degradation. The team\'s analysis revealed that '
        '"N+1 queries" in the reporting module accounted for 40% of all database '
        'load. After implementing "eager loading" and "query batching," average '
        'response times dropped from "850ms" to "120ms."'
    )

    # --- Section 3: Team Updates ---
    doc.add_heading("Team Updates", level=1)
    doc.add_paragraph(
        'The engineering department welcomed three new members this quarter. '
        'Elena Rodriguez joined as a "Senior DevOps Engineer," bringing '
        'expertise in "Kubernetes orchestration" and "GitOps workflows." '
        'David Park was hired as a "Staff Security Engineer" to lead the '
        '"zero-trust architecture" initiative.'
    )
    doc.add_paragraph(
        'Team lead Rachel Kim said: "Our hiring strategy focuses on finding '
        'people who aren\'t just technically skilled, but who also embrace '
        'our culture of \'continuous improvement.\' We don\'t just want coders; '
        'we want \'problem solvers.\'"'
    )

    # --- Section 4: Project Roadmap ---
    doc.add_heading("Project Roadmap", level=1)

    # Add a table for the roadmap
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"

    headers = ["Milestone", "Target Date", "Owner", "Status"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    roadmap_data = [
        ['"Phase 3" Launch', 'April 15, 2025', 'Sarah Chen', '"On Track"'],
        ['"Cloud Migration"', 'May 30, 2025', 'Elena Rodriguez', '"In Progress"'],
        ['"Mobile App" v2.0', 'June 20, 2025', 'David Park', '"Planning"'],
        ['"Analytics Dashboard"', 'July 10, 2025', 'Marcus Johnson', '"Backlog"'],
        ['"Security Audit" Phase 3', 'August 5, 2025', 'Rachel Kim', '"Scheduled"'],
    ]
    for r, row_data in enumerate(roadmap_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Section 5: Challenges and Lessons Learned ---
    doc.add_heading("Challenges and Lessons Learned", level=1)
    doc.add_paragraph(
        'The quarter wasn\'t without its challenges. The "production incident" '
        'on February 12th exposed weaknesses in our "monitoring pipeline." '
        'The post-mortem revealed that the root cause was a "race condition" '
        'in the caching layer, which the team described as a "known unknown" '
        'that had been deprioritized.'
    )
    doc.add_paragraph(
        'Lessons learned include: "Always test under realistic load conditions," '
        '"Don\'t skip integration tests to save time," and "Invest in '
        'observability before you need it." As CTO James Wright put it: '
        '"The best time to build monitoring was last quarter. The second best '
        'time is now."'
    )

    # --- Section 6: Conclusion ---
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        'Q1 2025 has been a "transformative quarter" for the Horizon project. '
        'The team\'s commitment to "engineering excellence" and "continuous '
        'delivery" has positioned us well for the challenges ahead. We\'re '
        'confident that the upcoming "Phase 3" launch will demonstrate the '
        'value of our investments in "scalability" and "reliability."'
    )
    doc.add_paragraph(
        'Next quarter\'s priorities include completing the "cloud migration," '
        'launching the "mobile app" redesign, and expanding the "automated '
        'testing" coverage to 95%. The team\'s shared vision is clear: '
        '"Build systems that don\'t just work today, but that scale for tomorrow."'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
