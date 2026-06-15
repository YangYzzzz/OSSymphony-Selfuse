"""
Reward Script: Employee Benefits Enrollment Guide
Task ID: writer_hr_081
Domain: libreoffice_writer
Scoring:
  Component 1: 7+ tables present (0.30)
  Component 2: Colored section dividers (0.20)
  Component 3: FAQ Q/A pair formatting with bold questions (0.15)
  Component 4: Medical/Dental/Vision as separate comparison tables (0.20)
  Component 5: Cost calculator table with Single/Couple/Family columns (0.15)
"""

import os

from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_081'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            import time
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def count_colored_dividers(doc):
    """Count paragraphs with colored background shading (non-white, non-auto fill)."""
    count = 0
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            shd = pPr.find(qn('w:shd'))
            if shd is not None:
                fill = shd.get(qn('w:fill'))
                if fill and fill.lower() not in ('auto', 'ffffff', '000000'):
                    count += 1
    return count


def find_table_with_header_keywords(doc, keywords):
    """Find a table whose first row contains all specified keywords (case-insensitive)."""
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header_text = ' '.join(cell.text.lower() for cell in table.rows[0].cells)
        if all(kw.lower() in header_text for kw in keywords):
            return table
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: At least 7 tables present (0.30 points)
    # Initial env has 0 tables; golden env has 8 tables.
    try:
        num_tables = len(doc.tables)
        if num_tables >= 7:
            print(f"PASS: Component 1 — {num_tables} tables found (>= 7 required) (0.30 pts)")
            total_score += 0.30
        elif num_tables >= 4:
            partial = round(0.30 * (num_tables / 7), 2)
            print(f"PARTIAL: Component 1 — {num_tables} tables found (need 7+), awarding {partial} pts")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — only {num_tables} tables found, need at least 7")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Colored section dividers — shaded paragraphs (0.20 points)
    # Initial env has 0 colored dividers; golden env has 10.
    try:
        divider_count = count_colored_dividers(doc)
        if divider_count >= 5:
            print(f"PASS: Component 2 — {divider_count} colored dividers found (>= 5 required) (0.20 pts)")
            total_score += 0.20
        elif divider_count >= 2:
            partial = round(0.20 * (divider_count / 5), 2)
            print(f"PARTIAL: Component 2 — {divider_count} colored dividers found, awarding {partial} pts")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — only {divider_count} colored dividers found, need at least 5")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: FAQ section with Q/A formatting and bold questions (0.15 points)
    # Initial env has no Q:/A: pairs; golden env has 6 Q/A pairs with bold questions.
    try:
        qa_pairs = 0
        bold_questions = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith('Q:') or text.startswith('Q.'):
                qa_pairs += 1
                # Check if the question text is bold
                for run in para.runs:
                    if run.text.strip() and run.font.bold:
                        bold_questions += 1
                        break

        if qa_pairs >= 3 and bold_questions >= 3:
            print(f"PASS: Component 3 — {qa_pairs} Q/A pairs, {bold_questions} bold questions (0.15 pts)")
            total_score += 0.15
        elif qa_pairs >= 2:
            partial = round(0.15 * min(qa_pairs, 3) / 3, 2)
            print(f"PARTIAL: Component 3 — {qa_pairs} Q/A pairs, {bold_questions} bold, awarding {partial} pts")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — {qa_pairs} Q/A pairs found, need at least 3 with bold formatting")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Separate Medical, Dental, Vision comparison tables (0.20 points)
    # Initial env has no tables at all; golden env has distinct comparison tables for each.
    try:
        tables_found = 0

        # Medical comparison table: should have columns like Bronze/Silver/Gold or plan names
        medical_table = find_table_with_header_keywords(doc, ['bronze'])
        if medical_table is None:
            medical_table = find_table_with_header_keywords(doc, ['feature'])
            # Check if it's actually a medical table (has deductible rows)
            if medical_table:
                all_text = ' '.join(cell.text.lower() for row in medical_table.rows for cell in row.cells)
                if 'deductible' not in all_text:
                    medical_table = None
        if medical_table:
            tables_found += 1
            print(f"  Medical comparison table found ({len(medical_table.rows)} rows)")

        # Dental comparison table
        dental_table = find_table_with_header_keywords(doc, ['dental'])
        if dental_table is None:
            # Look for table with preventive/basic/major procedure terminology
            for table in doc.tables:
                all_text = ' '.join(cell.text.lower() for row in table.rows for cell in row.cells)
                if 'preventive' in all_text and 'basic' in all_text and table != medical_table:
                    dental_table = table
                    break
        if dental_table and dental_table != medical_table:
            tables_found += 1
            print(f"  Dental comparison table found ({len(dental_table.rows)} rows)")

        # Vision comparison table
        vision_table = find_table_with_header_keywords(doc, ['vision'])
        if vision_table is None:
            for table in doc.tables:
                all_text = ' '.join(cell.text.lower() for row in table.rows for cell in row.cells)
                if ('eye exam' in all_text or 'frames' in all_text or 'contact' in all_text) and table != medical_table and table != dental_table:
                    vision_table = table
                    break
        if vision_table and vision_table != medical_table and vision_table != dental_table:
            tables_found += 1
            print(f"  Vision comparison table found ({len(vision_table.rows)} rows)")

        if tables_found >= 3:
            print(f"PASS: Component 4 — all 3 comparison tables found (medical, dental, vision) (0.20 pts)")
            total_score += 0.20
        elif tables_found >= 1:
            partial = round(0.20 * tables_found / 3, 2)
            print(f"PARTIAL: Component 4 — {tables_found}/3 comparison tables found, awarding {partial} pts")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — no medical/dental/vision comparison tables found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Cost calculator table with Single/Couple/Family columns (0.15 points)
    # Initial env has no tables; golden env has a deduction table with these coverage tiers.
    try:
        cost_table = find_table_with_header_keywords(doc, ['single', 'couple', 'family'])
        if cost_table is None:
            cost_table = find_table_with_header_keywords(doc, ['single', 'family'])

        if cost_table and len(cost_table.rows) >= 3:
            # Verify it has actual dollar amounts in the data rows
            dollar_count = 0
            for row in cost_table.rows[1:]:
                for cell in row.cells:
                    if '$' in cell.text:
                        dollar_count += 1
            if dollar_count >= 3:
                print(f"PASS: Component 5 — Cost calculator table found with {dollar_count} dollar values (0.15 pts)")
                total_score += 0.15
            elif dollar_count >= 1:
                print(f"PARTIAL: Component 5 — Cost table found but only {dollar_count} dollar values (0.05 pts)")
                total_score += 0.05
            else:
                print(f"FAIL: Component 5 — Cost table found but no dollar values")
        else:
            print(f"FAIL: Component 5 — No cost calculator table with Single/Couple/Family columns found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
