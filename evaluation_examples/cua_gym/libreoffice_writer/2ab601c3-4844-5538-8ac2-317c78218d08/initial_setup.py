"""
Initial Setup: Employee Benefits Enrollment Guide - raw text form
Task ID: writer_hr_081
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_081'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Title
    doc.add_heading('Open_Enrollment_Guide_2026', level=0)

    # --- Raw benefits data in paragraph form (NO tables, NO dividers) ---

    doc.add_heading('Medical Plans', level=1)
    doc.add_paragraph(
        'Apex Health offers three medical plan options for the 2026 benefits year. '
        'The Bronze Essential plan has a $3,500 individual deductible and $7,000 family deductible '
        'with 70/30 coinsurance after deductible. The monthly premium for single coverage is $185, '
        'couple coverage is $420, and family coverage is $610. Copay for primary care is $40 and '
        'specialist visits are $75. Emergency room visits have a $350 copay. '
        'Prescription drug tiers are $15 generic, $45 preferred brand, and $80 non-preferred brand.'
    )
    doc.add_paragraph(
        'The Silver Advantage plan features a $1,500 individual deductible and $3,000 family deductible '
        'with 80/20 coinsurance. Monthly premiums are $310 for single, $685 for couple, and $940 for family. '
        'Primary care copay is $25, specialist is $50, and ER is $250. '
        'Prescription tiers are $10 generic, $35 preferred brand, and $65 non-preferred brand.'
    )
    doc.add_paragraph(
        'The Gold Premier plan has a $500 individual deductible and $1,000 family deductible '
        'with 90/10 coinsurance. Monthly premiums are $475 for single, $1,020 for couple, '
        'and $1,385 for family. Primary care copay is $15, specialist is $30, and ER is $150. '
        'Prescription tiers are $5 generic, $20 preferred brand, and $45 non-preferred brand.'
    )

    doc.add_heading('Dental Plans', level=1)
    doc.add_paragraph(
        'Two dental plan options are available. The Basic Dental plan covers preventive care at 100%, '
        'basic procedures at 80%, and major procedures at 50%. The annual maximum benefit is $1,500 '
        'per person. Orthodontic coverage is not included. Monthly premiums are $28 single, $62 couple, '
        'and $95 family. The deductible is $50 individual and $150 family.'
    )
    doc.add_paragraph(
        'The Enhanced Dental plan covers preventive at 100%, basic at 90%, and major at 70%. '
        'The annual maximum is $3,000 per person. Orthodontic coverage is included at 50% '
        'with a $2,000 lifetime maximum for dependents under 26. Monthly premiums are $52 single, '
        '$108 couple, and $165 family. Deductible is $25 individual and $75 family.'
    )

    doc.add_heading('Vision Plans', level=1)
    doc.add_paragraph(
        'The Standard Vision plan includes an annual eye exam with a $10 copay. '
        'Frames allowance is $150 every 24 months. Contact lens allowance is $130 annually. '
        'Monthly premiums are $8 single, $16 couple, and $24 family.'
    )
    doc.add_paragraph(
        'The Premium Vision plan includes an annual eye exam with no copay. '
        'Frames allowance is $250 every 12 months. Contact lens allowance is $200 annually. '
        'Lens enhancements (progressive, anti-reflective) are covered at 80%. '
        'Monthly premiums are $15 single, $30 couple, and $45 family.'
    )

    doc.add_heading('Paycheck Deduction Information', level=1)
    doc.add_paragraph(
        'For employees paid biweekly (26 pay periods), the per-paycheck deductions are as follows. '
        'Bronze Essential medical: $92.50 single, $210.00 couple, $305.00 family. '
        'Silver Advantage medical: $155.00 single, $342.50 couple, $470.00 family. '
        'Gold Premier medical: $237.50 single, $510.00 couple, $692.50 family. '
        'Basic Dental: $14.00 single, $31.00 couple, $47.50 family. '
        'Enhanced Dental: $26.00 single, $54.00 couple, $82.50 family. '
        'Standard Vision: $4.00 single, $8.00 couple, $12.00 family. '
        'Premium Vision: $7.50 single, $15.00 couple, $22.50 family.'
    )

    doc.add_heading('Enrollment Timeline', level=1)
    doc.add_paragraph(
        'The open enrollment period for 2026 begins on October 15, 2025 and ends on November 14, 2025. '
        'Benefits orientation webinars will be held on October 8-9, 2025. '
        'The enrollment system opens on October 15, 2025 at 8:00 AM EST. '
        'One-on-one benefits counseling sessions are available October 20-31, 2025. '
        'The deadline for all enrollment changes is November 14, 2025 at 11:59 PM EST. '
        'Confirmation statements will be mailed by December 1, 2025. '
        'New benefits take effect on January 1, 2026.'
    )

    doc.add_heading('Provider Network Summary', level=1)
    doc.add_paragraph(
        'The Apex Health network includes approximately 45,000 physicians and 380 hospitals nationwide. '
        'In-network primary care physicians number around 18,500, with 12,200 specialists. '
        'There are 14,300 dentists in the dental network and 8,100 optometrists and ophthalmologists '
        'in the vision network. Urgent care facilities number 2,400 across all 50 states. '
        'The telemedicine platform is available 24/7 with board-certified physicians. '
        'Out-of-network coverage is available at reduced benefit levels on the Silver and Gold medical plans.'
    )

    doc.add_heading('FSA and HSA Information', level=1)
    doc.add_paragraph(
        'A Flexible Spending Account (FSA) is available with all medical plans. The 2026 contribution '
        'limit is $3,200. Funds must be used by March 15, 2027 (grace period). The employer contributes '
        'no match. A Dependent Care FSA is also available with a $5,000 annual limit. '
        'A Health Savings Account (HSA) is available only with the Bronze Essential plan (HDHP-qualified). '
        'The 2026 contribution limit is $4,300 individual and $8,550 family. Funds roll over indefinitely. '
        'The employer contributes $500 individual or $1,000 family annually. '
        'HSA funds can be invested once the balance exceeds $1,000.'
    )

    doc.add_heading('Life Insurance', level=1)
    doc.add_paragraph(
        'Basic life insurance at 1x annual salary (up to $200,000) is provided at no cost to employees. '
        'Supplemental Tier 1 provides 2x annual salary coverage up to $400,000 for $0.15 per $1,000 '
        'of coverage per month. Supplemental Tier 2 provides 3x annual salary up to $600,000 for $0.28 '
        'per $1,000 per month. Supplemental Tier 3 provides 4x annual salary up to $800,000 for $0.42 '
        'per $1,000 per month. Accidental death and dismemberment (AD&D) coverage matches the life '
        'insurance amount at each tier. Spouse and dependent life insurance is also available. '
        'Evidence of insurability is required for Tier 2 and above.'
    )

    doc.add_heading('Frequently Asked Questions', level=1)
    doc.add_paragraph(
        'Can I change my benefits outside of open enrollment? Generally no, unless you experience a '
        'qualifying life event such as marriage, birth of a child, divorce, or loss of other coverage. '
        'You must notify HR within 30 days of the qualifying event.'
    )
    doc.add_paragraph(
        'What happens if I do not enroll during open enrollment? If you do not make an active election, '
        'your current benefits will roll over to the next year. However, FSA elections do NOT roll over '
        'and you must re-enroll each year.'
    )
    doc.add_paragraph(
        'When does coverage begin for new hires? Benefits coverage begins on the first day of the month '
        'following 30 days of employment. New hires have 31 days from their start date to complete enrollment.'
    )
    doc.add_paragraph(
        'Can I cover my domestic partner? Yes, domestic partners and their eligible dependents may be '
        'enrolled in medical, dental, and vision plans. An affidavit of domestic partnership is required.'
    )
    doc.add_paragraph(
        'Are preventive services covered at 100%? Yes, all medical plans cover in-network preventive '
        'services at 100% with no cost sharing, as required by the Affordable Care Act. This includes '
        'annual physicals, immunizations, and age-appropriate screenings.'
    )
    doc.add_paragraph(
        'How do I find an in-network provider? Visit the Apex Health provider directory at '
        'providers.apexhealth.com or call Member Services at 1-800-555-APEX (2739). '
        'You can search by specialty, location, and language spoken.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
