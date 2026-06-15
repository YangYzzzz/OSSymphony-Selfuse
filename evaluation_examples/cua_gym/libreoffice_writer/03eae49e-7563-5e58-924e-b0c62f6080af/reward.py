"""
Reward Script: Create a #10 business envelope with delivery and return addresses
Task ID: writer_lec_037
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Envelope section exists (2+ sections)
  Component 2 (0.3): Envelope dimensions match #10 standard (9.5 x 4.125 inches)
  Component 3 (0.2): Return address present and left-aligned in envelope section
  Component 4 (0.2): Delivery address present and centered in envelope section
"""

import os
from docx import Document
from docx.shared import Inches, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_037'

# Tolerance for dimension checks: 0.05 inches
DIM_TOLERANCE = 0.05


def get_envelope_section(doc):
    """Return the envelope section (non-letter section) if it exists, else None."""
    if len(doc.sections) < 2:
        return None
    # The envelope section is additional to the original letter section.
    # Check each section beyond the first for envelope-like dimensions.
    for i in range(1, len(doc.sections)):
        s = doc.sections[i]
        w_in = s.page_width / 914400.0
        h_in = s.page_height / 914400.0
        # #10 envelope: 9.5 x 4.125 — could be either orientation
        if (abs(w_in - 9.5) < 0.5 and abs(h_in - 4.125) < 0.5) or \
           (abs(h_in - 9.5) < 0.5 and abs(w_in - 4.125) < 0.5):
            return i
    # If no exact match, return the last section as candidate
    return len(doc.sections) - 1


def get_section_paragraphs(doc, section_index):
    """Get paragraphs belonging to a specific section."""
    from docx.oxml.ns import qn
    body = doc.element.body

    # Collect all paragraph elements in order
    all_paras = list(body.iterchildren(qn('w:p')))

    # Find section break positions
    # Each paragraph may contain a sectPr (marking end of a section)
    section_breaks = []
    for idx, p_elem in enumerate(all_paras):
        ppr = p_elem.find(qn('w:pPr'))
        if ppr is not None:
            sect_pr = ppr.find(qn('w:sectPr'))
            if sect_pr is not None:
                section_breaks.append(idx)

    # Paragraphs are divided by section breaks
    # Section 0: paras[0..break0], Section 1: paras[break0+1..break1], etc.
    # Last section: paras after last break to end (its sectPr is in body, not in a para)
    starts = [0] + [b + 1 for b in section_breaks]
    ends = section_breaks + [len(all_paras)]

    if section_index >= len(starts):
        return []

    start = starts[section_index]
    end = ends[section_index]

    # Map element objects to doc.paragraphs by matching elements
    para_map = {p._element: p for p in doc.paragraphs}
    result = []
    for i in range(start, end):
        if all_paras[i] in para_map:
            result.append(para_map[all_paras[i]])
    return result


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Envelope section exists — document has 2+ sections (0.3 points)
    try:
        num_sections = len(doc.sections)
        if num_sections >= 2:
            print(f"PASS: Component 1 — Document has {num_sections} sections (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — Document has only {num_sections} section(s), expected 2+")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Envelope dimensions match #10 standard: 9.5 x 4.125 inches (0.3 points)
    try:
        env_idx = get_envelope_section(doc)
        if env_idx is not None and env_idx >= 1:
            s = doc.sections[env_idx]
            w_in = s.page_width / 914400.0
            h_in = s.page_height / 914400.0
            # Check both possible orientations
            match_normal = abs(w_in - 9.5) < DIM_TOLERANCE and abs(h_in - 4.125) < DIM_TOLERANCE
            match_rotated = abs(h_in - 9.5) < DIM_TOLERANCE and abs(w_in - 4.125) < DIM_TOLERANCE
            if match_normal or match_rotated:
                print(f"PASS: Component 2 — Envelope dimensions {w_in:.3f} x {h_in:.3f} inches match #10 (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Envelope dimensions {w_in:.3f} x {h_in:.3f} inches, expected 9.5 x 4.125")
        else:
            print("FAIL: Component 2 — No envelope section found to check dimensions")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Return address in upper-left (LEFT aligned) with correct text (0.2 points)
    try:
        env_idx = get_envelope_section(doc)
        if env_idx is not None and env_idx >= 1:
            env_paras = get_section_paragraphs(doc, env_idx)
            # Look for return address keywords in non-empty paragraphs
            env_texts = [(p.text.strip(), p.paragraph_format.alignment) for p in env_paras if p.text.strip()]

            return_keywords = ['ABC Consulting', 'Denver', '80201']
            left_aligned_texts = []
            for text, align in env_texts:
                # LEFT alignment is 0 or None (inherited default is LEFT)
                if align == WD_PARAGRAPH_ALIGNMENT.LEFT or align is None:
                    left_aligned_texts.append(text)

            # Check that return address keywords appear in left-aligned paragraphs
            joined_left = ' '.join(left_aligned_texts)
            matched_keywords = [kw for kw in return_keywords if kw in joined_left]

            if len(matched_keywords) >= 2:
                print(f"PASS: Component 3 — Return address found left-aligned with keywords: {matched_keywords} (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — Return address not found. Left-aligned texts: {left_aligned_texts}")
        else:
            print("FAIL: Component 3 — No envelope section found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Delivery address centered with correct text (0.2 points)
    try:
        env_idx = get_envelope_section(doc)
        if env_idx is not None and env_idx >= 1:
            env_paras = get_section_paragraphs(doc, env_idx)
            env_texts = [(p.text.strip(), p.paragraph_format.alignment) for p in env_paras if p.text.strip()]

            delivery_keywords = ['Dr. Emily Chen', 'Chicago', '60601']
            centered_texts = []
            for text, align in env_texts:
                if align == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    centered_texts.append(text)

            joined_center = ' '.join(centered_texts)
            matched_keywords = [kw for kw in delivery_keywords if kw in joined_center]

            if len(matched_keywords) >= 2:
                print(f"PASS: Component 4 — Delivery address found centered with keywords: {matched_keywords} (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 4 — Delivery address not properly centered. Centered texts: {centered_texts}")
        else:
            print("FAIL: Component 4 — No envelope section found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point — persist app state then verify
def persist_app_state(domain):
    import os, time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
