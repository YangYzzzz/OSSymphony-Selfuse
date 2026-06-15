"""
Initial Setup: Create a letter document addressed to Dr. Emily Chen
Task ID: writer_lec_037
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_037'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Standard letter page setup (8.5 x 11 inches)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Sender information
    sender_para = doc.add_paragraph()
    sender_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sender_para.paragraph_format.space_after = Pt(0)
    run = sender_para.add_run("ABC Consulting")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    for line in ["100 Main St, Suite 200", "Denver, CO 80201"]:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"

    # Date
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(24)
    date_para.paragraph_format.space_after = Pt(12)
    r = date_para.add_run("March 15, 2025")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    # Recipient address block
    recipient_lines = [
        "Dr. Emily Chen",
        "Global Health Institute",
        "789 Medical Center Dr",
        "Chicago, IL 60601",
    ]
    for line in recipient_lines:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"

    # Salutation
    sal = doc.add_paragraph()
    sal.paragraph_format.space_before = Pt(12)
    sal.paragraph_format.space_after = Pt(12)
    r = sal.add_run("Dear Dr. Chen,")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    # Letter body
    body_text = (
        "Thank you for your recent inquiry regarding our strategic consulting "
        "services for healthcare organizations. We appreciate the opportunity to "
        "discuss how ABC Consulting can support the Global Health Institute's "
        "upcoming research initiative on community health outcomes."
    )
    body = doc.add_paragraph()
    body.paragraph_format.space_after = Pt(12)
    body.paragraph_format.first_line_indent = Inches(0.5)
    r = body.add_run(body_text)
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    body_text2 = (
        "Our team has extensive experience working with medical research "
        "institutions, and we believe our data analytics and project management "
        "expertise would be a valuable addition to your research framework. "
        "We would be delighted to schedule a meeting at your earliest convenience "
        "to discuss the project scope and timeline in more detail."
    )
    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after = Pt(12)
    body2.paragraph_format.first_line_indent = Inches(0.5)
    r = body2.add_run(body_text2)
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    body_text3 = (
        "Please do not hesitate to reach out if you require any additional "
        "information or would like to arrange a preliminary call with our "
        "senior consultants. We look forward to the possibility of collaborating "
        "with the Global Health Institute."
    )
    body3 = doc.add_paragraph()
    body3.paragraph_format.space_after = Pt(24)
    body3.paragraph_format.first_line_indent = Inches(0.5)
    r = body3.add_run(body_text3)
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    # Closing
    closing = doc.add_paragraph()
    closing.paragraph_format.space_after = Pt(0)
    r = closing.add_run("Sincerely,")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    # Signature space + name
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(36)
    sig.paragraph_format.space_after = Pt(0)
    r = sig.add_run("James Parker")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    title_line = doc.add_paragraph()
    title_line.paragraph_format.space_after = Pt(0)
    title_line.paragraph_format.space_before = Pt(0)
    r = title_line.add_run("Senior Consultant")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    company_line = doc.add_paragraph()
    company_line.paragraph_format.space_before = Pt(0)
    r = company_line.add_run("ABC Consulting")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
