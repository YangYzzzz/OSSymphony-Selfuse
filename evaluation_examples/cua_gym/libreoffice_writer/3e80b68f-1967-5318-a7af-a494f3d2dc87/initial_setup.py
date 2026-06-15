"""
initial_setup.py - Create a Writer document with 4 chapters, each starting with
a heading and a plain summary paragraph followed by body paragraphs.
No borders or special formatting on summary paragraphs.
"""

import subprocess
import shlex
import os
import time

WORKDIR = '/home/user'
FILEPATH = os.path.join(WORKDIR, 'writer_tech_044.docx')

# Install python-docx
subprocess.run(['pip3', 'install', 'python-docx'], check=True,
               stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Page setup
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

chapters = [
    {
        "title": "Chapter 1: Introduction to Machine Learning",
        "summary": "This chapter provides an overview of machine learning concepts, including supervised and unsupervised learning paradigms, and sets the foundation for understanding modern AI systems.",
        "body": [
            "Machine learning has emerged as one of the most transformative technologies of the 21st century. It enables computers to learn from data without being explicitly programmed, opening up possibilities that were previously unimaginable.",
            "The field traces its roots back to the 1950s, when Arthur Samuel coined the term while working on a checkers-playing program at IBM. Since then, the discipline has evolved dramatically, driven by increases in computing power and the availability of large datasets.",
            "In supervised learning, the algorithm is trained on labeled examples, where both the input and the desired output are provided. Common applications include image classification, spam detection, and medical diagnosis.",
            "Unsupervised learning, on the other hand, works with unlabeled data. The algorithm must discover hidden patterns and structures on its own. Clustering and dimensionality reduction are typical unsupervised techniques."
        ]
    },
    {
        "title": "Chapter 2: Neural Networks and Deep Learning",
        "summary": "This chapter explores the architecture of neural networks, from simple perceptrons to deep learning models, covering activation functions, backpropagation, and training strategies.",
        "body": [
            "Neural networks are computational models inspired by the biological neural networks in the human brain. They consist of interconnected nodes organized in layers that process information in a hierarchical manner.",
            "The simplest form of a neural network is the perceptron, a single-layer model that can only solve linearly separable problems. By stacking multiple layers, we create multi-layer perceptrons capable of learning complex nonlinear relationships.",
            "Deep learning refers to neural networks with many hidden layers. These deep architectures have achieved state-of-the-art results in image recognition, natural language processing, speech recognition, and many other domains.",
            "Training deep networks requires careful initialization, appropriate learning rates, and regularization techniques to prevent overfitting. Batch normalization and dropout are widely used strategies to improve training stability."
        ]
    },
    {
        "title": "Chapter 3: Natural Language Processing",
        "summary": "This chapter covers the fundamentals of natural language processing, including tokenization, word embeddings, transformer architectures, and large language models.",
        "body": [
            "Natural language processing bridges the gap between human communication and computer understanding. It encompasses a wide range of tasks from basic text classification to complex dialogue systems and machine translation.",
            "Tokenization is the first step in most NLP pipelines, breaking text into meaningful units. Modern subword tokenization methods like BPE and WordPiece balance vocabulary size with the ability to handle out-of-vocabulary words.",
            "Word embeddings represent words as dense vectors in a continuous space, capturing semantic relationships. Word2Vec and GloVe were pioneering approaches, while contextual embeddings from models like BERT represent the current state of the art.",
            "The transformer architecture, introduced in the landmark paper 'Attention Is All You Need,' revolutionized NLP by enabling parallel processing of sequences and capturing long-range dependencies through self-attention mechanisms."
        ]
    },
    {
        "title": "Chapter 4: Reinforcement Learning",
        "summary": "This chapter introduces reinforcement learning concepts including agents, environments, reward signals, and policy optimization methods used in game playing and robotics.",
        "body": [
            "Reinforcement learning is a paradigm where an agent learns to make decisions by interacting with an environment. The agent receives rewards or penalties based on its actions and aims to maximize cumulative reward over time.",
            "The exploration-exploitation tradeoff is a fundamental challenge in reinforcement learning. The agent must balance trying new actions to discover potentially better strategies with exploiting known good actions to accumulate reward.",
            "Q-learning and SARSA are classic value-based methods that estimate the expected return for each state-action pair. Deep Q-Networks extended these ideas to high-dimensional state spaces by using neural networks as function approximators.",
            "Policy gradient methods directly optimize the policy function without maintaining value estimates. Algorithms like REINFORCE, PPO, and A3C have achieved remarkable success in complex environments including Atari games and robotic control tasks."
        ]
    }
]

for i, ch in enumerate(chapters):
    if i > 0:
        doc.add_page_break()

    # Chapter heading (Heading 1)
    doc.add_heading(ch["title"], level=1)

    # Summary paragraph - plain Normal style, no borders
    summary_para = doc.add_paragraph(ch["summary"])
    summary_para.style = doc.styles['Normal']

    # Body paragraphs
    for body_text in ch["body"]:
        para = doc.add_paragraph(body_text)
        para.style = doc.styles['Normal']

doc.save(FILEPATH)
print(f"Document saved to {FILEPATH}")

# Launch LibreOffice Writer with the document
def launch_gui(command: str, delay_sec: float = 1.0):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)

launch_gui(f'libreoffice --writer "{FILEPATH}"', delay_sec=2.0)
print("LibreOffice Writer launched.")
