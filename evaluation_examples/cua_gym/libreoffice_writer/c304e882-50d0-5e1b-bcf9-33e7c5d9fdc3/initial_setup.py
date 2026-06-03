"""
Initial Setup: Behavioral Interview Questionnaire template (pre-task state)
Task ID: writer_hr_029
Domain: libreoffice_writer

The initial document contains only the title and candidate information section.
The agent must add the 10 behavioral questions, notes lines, and scoring section.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_029'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading('Behavioral Interview Questionnaire', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Candidate Information Section ---
    doc.add_paragraph('')  # spacer

    info_heading = doc.add_heading('Candidate Information', level=2)

    # Name field
    p_name = doc.add_paragraph()
    run_name = p_name.add_run('Name: ')
    run_name.bold = True
    run_name.font.size = Pt(12)
    p_name.add_run('_' * 40)

    # Position field
    p_pos = doc.add_paragraph()
    run_pos = p_pos.add_run('Position Applied For: ')
    run_pos.bold = True
    run_pos.font.size = Pt(12)
    p_pos.add_run('_' * 30)

    # Date field
    p_date = doc.add_paragraph()
    run_date = p_date.add_run('Interview Date: ')
    run_date.bold = True
    run_date.font.size = Pt(12)
    p_date.add_run('_' * 30)

    # Interviewer field
    p_int = doc.add_paragraph()
    run_int = p_int.add_run('Interviewer: ')
    run_int.bold = True
    run_int.font.size = Pt(12)
    p_int.add_run('_' * 35)

    doc.add_paragraph('')  # spacer

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
