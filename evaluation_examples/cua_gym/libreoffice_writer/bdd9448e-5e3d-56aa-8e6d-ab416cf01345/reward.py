"""
Reward Script: Avery 5167 return address labels (80 per sheet)
Task ID: writer_lec_042
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Table exists with ~20 rows x 4 cols (80 labels)
  Component 2 (0.5): All cells contain the correct return address
  Component 3 (0.2): Page margins adjusted for label layout (reduced from defaults)
"""

import os

from docx import Document
from docx.shared import Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_042'

# Expected address components
EXPECTED_NAME = "Maria Garcia"
EXPECTED_STREET = "555 Pine Road"
EXPECTED_CITY_STATE_ZIP = "Austin, TX 78701"

# Avery 5167: 4 columns, 20 rows = 80 labels
EXPECTED_COLS = 4
EXPECTED_ROWS = 20
EXPECTED_TOTAL = 80

# Default Writer margins (EMU): left=1143000, right=1143000
# Label layout should have smaller margins
DEFAULT_SIDE_MARGIN_EMU = 1143000


def persist_app_state(domain: str):
    """Save any unsaved document state via Ctrl+S."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table with correct grid dimensions (0.3 points)
    # Initial doc has 0 tables, golden has 1 table with 20x4
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 1 -- No tables found in document")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            total_cells = num_rows * num_cols

            if num_cols == EXPECTED_COLS and num_rows == EXPECTED_ROWS:
                print(f"PASS: Component 1 -- Table is {num_rows}x{num_cols} = {total_cells} labels (0.3 pts)")
                total_score += 0.3
            elif num_cols == EXPECTED_COLS and abs(num_rows - EXPECTED_ROWS) <= 2:
                # Close but not exact row count -- partial credit
                print(f"PARTIAL: Component 1 -- Table is {num_rows}x{num_cols}, expected {EXPECTED_ROWS}x{EXPECTED_COLS} (0.15 pts)")
                total_score += 0.15
            elif total_cells >= 60:
                # Has a substantial label grid but wrong dimensions
                print(f"PARTIAL: Component 1 -- Table has {total_cells} cells ({num_rows}x{num_cols}), expected 80 (0.1 pts)")
                total_score += 0.1
            else:
                print(f"FAIL: Component 1 -- Table is {num_rows}x{num_cols} = {total_cells}, expected {EXPECTED_ROWS}x{EXPECTED_COLS} = {EXPECTED_TOTAL}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: All cells contain the correct return address (0.5 points)
    # Initial doc has no tables so 0 correct cells; golden has 80/80
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 2 -- No tables, cannot check cell contents")
        else:
            table = doc.tables[0]
            correct_count = 0
            total_cells = 0

            for row in table.rows:
                for cell in row.cells:
                    total_cells += 1
                    text = cell.text.strip()
                    has_name = EXPECTED_NAME.lower() in text.lower()
                    has_street = EXPECTED_STREET.lower() in text.lower()
                    has_city = "austin" in text.lower() and "78701" in text

                    if has_name and has_street and has_city:
                        correct_count += 1

            if total_cells > 0:
                ratio = correct_count / max(total_cells, EXPECTED_TOTAL)
                # Scale: need at least 75% correct for any credit
                if ratio >= 0.95:
                    score = 0.5
                elif ratio >= 0.75:
                    score = round(0.5 * (ratio - 0.75) / 0.25, 2)
                else:
                    score = 0.0

                if score > 0:
                    print(f"PASS: Component 2 -- {correct_count}/{total_cells} cells have correct address ({score} pts)")
                    total_score += score
                else:
                    print(f"FAIL: Component 2 -- Only {correct_count}/{total_cells} cells have correct address (need >= 75%)")
            else:
                print("FAIL: Component 2 -- Table has 0 cells")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Page margins adjusted for label layout (0.2 points)
    # Initial doc has default margins (1143000 EMU = 1.25 in per side)
    # Golden has reduced margins (274320 EMU = 0.3 in per side)
    # Label sheets need tighter margins than default
    try:
        section = doc.sections[0]
        left_m = section.left_margin
        right_m = section.right_margin

        # Check that margins are significantly smaller than defaults
        # Default side margins are 1143000 EMU (1.25 in). Label layout needs less.
        if left_m < DEFAULT_SIDE_MARGIN_EMU * 0.7 and right_m < DEFAULT_SIDE_MARGIN_EMU * 0.7:
            print(f"PASS: Component 3 -- Margins adjusted: L={left_m/914400:.2f}in, R={right_m/914400:.2f}in (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 -- Margins not reduced for label layout: L={left_m/914400:.2f}in, R={right_m/914400:.2f}in (defaults ~1.25in)")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
