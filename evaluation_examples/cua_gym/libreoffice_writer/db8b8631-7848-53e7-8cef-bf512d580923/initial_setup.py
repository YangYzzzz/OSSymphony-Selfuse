"""
Initial Setup: Insert cross-reference to bookmark in Writer document
Task ID: writer_af_045
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_af_045'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_body_text(doc, text):
    p = doc.add_paragraph(text)
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    return p


def add_bookmark(paragraph, bookmark_name, text):
    """Add a bookmark around text in a paragraph."""
    # Create bookmarkStart
    bm_start = paragraph._element.makeelement(qn('w:bookmarkStart'), {
        qn('w:id'): '1',
        qn('w:name'): bookmark_name,
    })
    # Create the run with text
    run = paragraph.add_run(text)
    # Create bookmarkEnd
    bm_end = paragraph._element.makeelement(qn('w:bookmarkEnd'), {
        qn('w:id'): '1',
    })
    # Insert bookmarkStart before the run, bookmarkEnd after
    run_elem = run._element
    run_elem.addprevious(bm_start)
    run_elem.addnext(bm_end)
    return run


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ========== PAGE 1: Title Page ==========
    doc.add_paragraph()  # spacer
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('Research Findings Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph('Comprehensive Analysis of Operational Efficiency Improvements')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()
    author = doc.add_paragraph('Prepared by: Dr. Elena Rodriguez, Senior Research Analyst')
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_p = doc.add_paragraph('Date: March 15, 2025')
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    org = doc.add_paragraph('Meridian Research Institute')
    org.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_page_break(doc)

    # ========== PAGE 2: Table of Contents ==========
    add_heading_styled(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary ........................ 3',
        '2. Methodology .............................. 4',
        '3. Key Findings ............................. 4',
        '4. Statistical Analysis ..................... 5',
        '5. Market Comparison ........................ 6',
        '6. Implementation Recommendations ........... 7',
        '7. Risk Assessment .......................... 8',
        '8. Financial Projections .................... 9',
        '9. Cross-Reference Summary .................. 10',
        '10. Conclusion .............................. 11',
        '11. Appendix A: Raw Data .................... 12',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    add_page_break(doc)

    # ========== PAGE 3: Executive Summary ==========
    add_heading_styled(doc, '1. Executive Summary', level=1)
    add_body_text(doc, (
        'This report presents the findings of a twelve-month longitudinal study '
        'conducted across 47 manufacturing facilities in the Asia-Pacific region. '
        'The research was initiated in January 2024 following preliminary observations '
        'suggesting that targeted process modifications could yield measurable gains '
        'in operational throughput.'
    ))
    add_body_text(doc, (
        'Over the study period, participating facilities implemented a phased '
        'approach to workflow optimization. The first phase involved baseline '
        'measurements using standardized KPI frameworks. The second phase introduced '
        'controlled variables including revised scheduling algorithms, updated '
        'equipment calibration protocols, and enhanced quality control checkpoints.'
    ))
    add_body_text(doc, (
        'The aggregate data reveals statistically significant improvements across '
        'multiple dimensions. Production cycle times decreased by an average of '
        '11.3%, while defect rates fell from 2.8% to 1.6%. Employee satisfaction '
        'scores, measured via quarterly surveys, rose by 18 points on a 100-point scale. '
        'These outcomes collectively support the hypothesis that systematic process '
        'refinement, when properly managed, delivers compounding benefits.'
    ))
    add_body_text(doc, (
        'The financial impact was equally noteworthy. Facilities that fully adopted '
        'the recommended protocols reported an average cost reduction of $1.2 million '
        'annually, with the earliest adopters seeing returns within the first quarter. '
        'Total projected savings across all participants over a five-year horizon '
        'exceed $280 million.'
    ))

    add_page_break(doc)

    # ========== PAGE 4: Methodology & Key Findings ==========
    add_heading_styled(doc, '2. Methodology', level=1)
    add_body_text(doc, (
        'The study employed a mixed-methods research design combining quantitative '
        'performance metrics with qualitative assessments from facility managers. '
        'Data collection instruments included automated sensor arrays, daily production '
        'logs, and bi-weekly structured interviews with operations leads.'
    ))
    add_body_text(doc, (
        'Statistical validity was ensured through stratified random sampling '
        'across facility sizes (small: <200 employees, medium: 200-500, large: >500). '
        'Control groups comprised 12 facilities that maintained existing protocols '
        'throughout the study period.'
    ))

    add_heading_styled(doc, '3. Key Findings', level=1)

    # The bookmarked sentence
    p_bookmark = doc.add_paragraph()
    p_bookmark.paragraph_format.space_after = Pt(6)
    add_bookmark(p_bookmark, 'KeyFinding',
                 'The primary result shows a 15% improvement in efficiency')

    add_body_text(doc, (
        'This improvement was consistent across facility sizes, though large '
        'facilities (>500 employees) demonstrated the strongest gains at 17.2%. '
        'Medium facilities averaged 14.8%, while small facilities achieved 13.1%. '
        'The variance between groups was not statistically significant (p=0.087), '
        'suggesting the methodology is broadly applicable.'
    ))
    add_body_text(doc, (
        'Secondary findings include a 22% reduction in equipment downtime, '
        'a 9% improvement in raw material utilization, and a 31% decrease in '
        'compliance-related incidents. These ancillary benefits were not '
        'initially hypothesized but emerged consistently across the dataset.'
    ))

    add_page_break(doc)

    # ========== PAGE 5: Statistical Analysis ==========
    add_heading_styled(doc, '4. Statistical Analysis', level=1)
    add_body_text(doc, (
        'A multivariate regression analysis was conducted to isolate the '
        'contribution of each intervention component. The model achieved an '
        'adjusted R-squared of 0.847, indicating strong explanatory power. '
        'The scheduling algorithm modification alone accounted for 38% of the '
        'observed improvement, while calibration protocol changes contributed 29% '
        'and quality control enhancements explained 24%.'
    ))
    add_body_text(doc, (
        'Time series analysis revealed that improvements were not linear. '
        'Facilities typically experienced a 3-4 week adjustment period during '
        'which productivity temporarily declined by 2-5%. Following this '
        'adaptation phase, performance metrics showed exponential improvement '
        'before stabilizing at new baselines after approximately 10 weeks.'
    ))
    add_body_text(doc, (
        'Confidence intervals for the primary efficiency metric ranged from '
        '13.2% to 16.8% at the 95% confidence level. The Shapiro-Wilk test '
        'confirmed normal distribution of residuals (W=0.983, p=0.412), '
        'validating the regression assumptions.'
    ))
    add_body_text(doc, (
        'A Granger causality test confirmed that the intervention preceded '
        'the efficiency gains (F=14.73, p<0.001), strengthening the causal '
        'interpretation of the results. No significant autocorrelation was '
        'detected in the residuals (Durbin-Watson d=1.96).'
    ))

    add_page_break(doc)

    # ========== PAGE 6: Market Comparison ==========
    add_heading_styled(doc, '5. Market Comparison', level=1)
    add_body_text(doc, (
        'To contextualize these findings, we compared our results against '
        'industry benchmarks published by the Global Manufacturing Efficiency '
        'Council (GMEC). The average industry improvement rate for similar '
        'interventions over the past decade has been 8.3%, placing our 15% '
        'result well above the historical norm.'
    ))
    add_body_text(doc, (
        'Regional comparisons revealed that facilities in Southeast Asia '
        'outperformed those in East Asia by a margin of 2.1 percentage points. '
        'This difference was attributed to lower baseline efficiency in Southeast '
        'Asian facilities, providing greater room for improvement. European '
        'facilities in a parallel GMEC study achieved 10.7% improvement using '
        'comparable methodologies.'
    ))
    add_body_text(doc, (
        'Sector-specific analysis showed that automotive component manufacturers '
        'benefited most (18.4% improvement), followed by electronics assembly '
        '(15.9%) and consumer goods packaging (12.3%). Heavy industry and '
        'chemical processing facilities showed more modest gains of 9.8% and '
        '8.5% respectively, likely due to regulatory constraints limiting '
        'process modifications.'
    ))
    add_body_text(doc, (
        'A competitive landscape review indicates that early adopters of these '
        'methodologies could achieve 2-3 year advantages over peers. The '
        'compound effect of sustained 15% annual efficiency gains translates '
        'to a cumulative 74.9% improvement over five years, fundamentally '
        'altering competitive positioning.'
    ))

    add_page_break(doc)

    # ========== PAGE 7: Implementation Recommendations ==========
    add_heading_styled(doc, '6. Implementation Recommendations', level=1)
    add_body_text(doc, (
        'Based on the findings, we recommend a three-phase implementation '
        'strategy. Phase I (Months 1-3) focuses on baseline assessment and '
        'stakeholder alignment. During this phase, facilities should conduct '
        'comprehensive audits of existing workflows, identify bottlenecks, '
        'and establish measurement frameworks.'
    ))
    add_body_text(doc, (
        'Phase II (Months 4-9) involves controlled deployment of the three '
        'primary interventions: scheduling optimization, equipment calibration '
        'updates, and quality control process redesign. We recommend staggered '
        'rollout beginning with the scheduling component, as it demonstrated '
        'the highest individual contribution to overall improvement.'
    ))
    add_body_text(doc, (
        'Phase III (Months 10-12) centers on consolidation and continuous '
        'improvement. Facilities should establish permanent monitoring '
        'dashboards, train internal champions, and develop feedback loops '
        'to sustain and incrementally enhance the gains achieved.'
    ))
    add_body_text(doc, (
        'Critical success factors include executive sponsorship, dedicated '
        'project management resources (minimum 0.5 FTE per facility), and '
        'transparent communication with frontline workers. Facilities that '
        'lacked these elements in our study showed 40% lower improvement rates. '
        'Budget allocation should anticipate $45,000-$120,000 per facility '
        'depending on size, with ROI typically achieved within 6-8 months.'
    ))

    add_page_break(doc)

    # ========== PAGE 8: Risk Assessment ==========
    add_heading_styled(doc, '7. Risk Assessment', level=1)
    add_body_text(doc, (
        'While the results are highly encouraging, several risk factors '
        'warrant consideration. The most significant risk is implementation '
        'fatigue, observed in 15% of participating facilities where initial '
        'enthusiasm waned during the adaptation period. Mitigation strategies '
        'include milestone celebrations, visible leadership engagement, and '
        'short-cycle feedback mechanisms.'
    ))
    add_body_text(doc, (
        'Technology dependency represents a moderate risk. The scheduling '
        'optimization component relies on real-time data feeds from IoT '
        'sensors. Facilities with aging infrastructure may require capital '
        'investment of $25,000-$75,000 to upgrade sensor networks before '
        'the methodology can be fully deployed.'
    ))
    add_body_text(doc, (
        'Market volatility could affect the relevance of optimization targets. '
        'If demand patterns shift significantly, the scheduling algorithms '
        'may require recalibration. We recommend quarterly reviews of '
        'optimization parameters against current market conditions.'
    ))
    add_body_text(doc, (
        'Regulatory changes in environmental compliance, particularly in '
        'the EU and ASEAN regions, may necessitate process modifications '
        'that could temporarily offset efficiency gains. Proactive monitoring '
        'of regulatory developments is strongly advised. The overall risk '
        'profile is assessed as moderate, with well-defined mitigation paths '
        'for all identified scenarios.'
    ))

    add_page_break(doc)

    # ========== PAGE 9: Financial Projections ==========
    add_heading_styled(doc, '8. Financial Projections', level=1)
    add_body_text(doc, (
        'The financial model projects cumulative savings of $280 million '
        'across all 47 facilities over a five-year period under the baseline '
        'scenario. The optimistic scenario, assuming 18% efficiency gains '
        'and accelerated adoption, projects $340 million. The conservative '
        'scenario, accounting for partial adoption and attrition, projects '
        '$195 million.'
    ))
    add_body_text(doc, (
        'Year-by-year projections for the baseline scenario are as follows: '
        'Year 1: $32 million (reflecting phased rollout), Year 2: $56 million, '
        'Year 3: $64 million, Year 4: $65 million, Year 5: $63 million. '
        'The slight decline in Years 4-5 reflects diminishing marginal returns '
        'as facilities approach optimal efficiency levels.'
    ))
    add_body_text(doc, (
        'Capital expenditure requirements total $4.7 million across all '
        'facilities, yielding a portfolio IRR of 287% and a payback period '
        'of 4.2 months. Even under the conservative scenario, the IRR exceeds '
        '150%, making this one of the highest-return operational investments '
        'available to the manufacturing sector.'
    ))
    add_body_text(doc, (
        'Sensitivity analysis reveals that the break-even efficiency improvement '
        'threshold is 4.8%. Given that all facilities in the study exceeded '
        'this threshold, the probability of negative returns is assessed at '
        'less than 2%. Monte Carlo simulation with 10,000 iterations confirmed '
        'a 95% confidence interval for total savings of $178-$362 million.'
    ))

    add_page_break(doc)

    # ========== PAGE 10: Cross-Reference Summary ==========
    add_heading_styled(doc, '9. Cross-Reference Summary', level=1)
    add_body_text(doc, (
        'This section consolidates the key data points referenced throughout '
        'the report. The purpose is to provide a quick-reference guide for '
        'executive decision-makers who may not read the full document.'
    ))
    add_body_text(doc, (
        'The efficiency improvement data has been validated through multiple '
        'analytical approaches including regression analysis, time series '
        'decomposition, and cross-sectional comparison. All methods converge '
        'on consistent conclusions.'
    ))
    add_body_text(doc, (
        'The competitive analysis in Section 5 demonstrates that these gains '
        'significantly exceed industry averages, providing a strong basis for '
        'strategic investment decisions. The financial projections in Section 8 '
        'further reinforce the business case with robust ROI estimates.'
    ))

    # The line where the cross-reference should be inserted
    # Initial state: just ends with "As described earlier: " and NO cross-reference
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.space_after = Pt(6)
    p_ref.add_run('As described earlier: ')

    add_body_text(doc, (
        'The above reference should dynamically update if the source text '
        'is modified, ensuring consistency across the document. This mechanism '
        'is essential for maintaining accuracy in living documents that undergo '
        'frequent revisions.'
    ))

    add_page_break(doc)

    # ========== PAGE 11: Conclusion ==========
    add_heading_styled(doc, '10. Conclusion', level=1)
    add_body_text(doc, (
        'The evidence presented in this report unequivocally supports the '
        'adoption of the proposed efficiency optimization methodology. The '
        '15% improvement in operational efficiency, validated across 47 '
        'facilities over twelve months, represents a transformative opportunity '
        'for the manufacturing sector.'
    ))
    add_body_text(doc, (
        'The methodology is scalable, replicable, and adaptable to diverse '
        'manufacturing contexts. The phased implementation approach minimizes '
        'disruption while maximizing the probability of successful adoption. '
        'The financial returns are exceptional by any standard of operational '
        'investment.'
    ))
    add_body_text(doc, (
        'We recommend immediate initiation of Phase I activities across '
        'all targeted facilities. The competitive window for early-mover '
        'advantage is estimated at 18-24 months, after which widespread '
        'adoption will normalize performance levels. Organizations that '
        'act decisively will secure lasting structural advantages in their '
        'respective markets.'
    ))
    add_body_text(doc, (
        'Further research is planned to explore the applicability of these '
        'findings to service industries and knowledge-work environments. '
        'Preliminary indicators suggest that analogous improvements of 10-12% '
        'may be achievable in these sectors, potentially expanding the total '
        'addressable impact by an order of magnitude.'
    ))

    add_page_break(doc)

    # ========== PAGE 12: Appendix A ==========
    add_heading_styled(doc, '11. Appendix A: Raw Data Summary', level=1)
    add_body_text(doc, (
        'The following table summarizes raw performance data from a '
        'representative sample of facilities included in the study.'
    ))

    # Add a data table
    table = doc.add_table(rows=11, cols=5)
    table.style = 'Table Grid'
    headers = ['Facility ID', 'Region', 'Size', 'Baseline Efficiency', 'Post-Intervention']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['MFG-001', 'Singapore', 'Large', '72.3%', '84.1%'],
        ['MFG-007', 'Vietnam', 'Medium', '68.9%', '79.5%'],
        ['MFG-012', 'Thailand', 'Small', '75.1%', '85.0%'],
        ['MFG-018', 'Malaysia', 'Large', '70.6%', '81.8%'],
        ['MFG-023', 'Indonesia', 'Medium', '66.4%', '76.9%'],
        ['MFG-029', 'Philippines', 'Small', '71.8%', '81.2%'],
        ['MFG-034', 'Japan', 'Large', '78.2%', '89.7%'],
        ['MFG-038', 'South Korea', 'Medium', '76.5%', '87.3%'],
        ['MFG-041', 'Taiwan', 'Large', '74.0%', '85.6%'],
        ['MFG-045', 'China', 'Small', '69.7%', '80.4%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    add_body_text(doc, '')
    add_body_text(doc, (
        'Note: Efficiency is measured as the ratio of actual output to '
        'theoretical maximum output, adjusted for planned downtime and '
        'scheduled maintenance windows. Baseline measurements represent '
        'the average of the three months preceding intervention deployment.'
    ))
    add_body_text(doc, (
        'Complete raw datasets are available upon request from the '
        'Meridian Research Institute data repository (reference: MRI-2025-EFF-047). '
        'All data has been anonymized in compliance with participant '
        'confidentiality agreements.'
    ))

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
