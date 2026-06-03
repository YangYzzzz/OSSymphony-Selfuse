"""
Reward Script: Add bibliography entry for Lee (2021) and insert citation in 4th paragraph
Task ID: osworld_writer_bibliography_crossref_006
Domain: libreoffice_writer
Scoring:
  Component 1: Lee (2021) bibliography entry added to bibliography section (0.5 pts)
  Component 2: Cross-reference/citation '(Lee, 2021)' inserted in the 4th paragraph (0.5 pts)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_bibliography_crossref_006'

def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # --- Component 1: Lee (2021) bibliography entry added (0.5 points) ---
    # The task requires adding 'Lee, K. (2021). Data Privacy in Healthcare. Medical Press.'
    # to the bibliography section. The bibliography section follows the 'Bibliography' heading.
    # We check that at least one paragraph contains 'Lee' and '2021' with key details.
    try:
        lee_entry_found = False
        in_bibliography = False
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading') and 'Bibliography' in para.text:
                in_bibliography = True
                continue
            if in_bibliography and para.style.name.startswith('Heading'):
                # End of bibliography section
                in_bibliography = False
            if in_bibliography:
                text = para.text.strip()
                # Check for the key components of the Lee entry
                if ('Lee' in text and '2021' in text and
                        ('Data Privacy' in text or 'Healthcare' in text or 'Medical Press' in text)):
                    lee_entry_found = True
                    print(f"PASS: Component 1 — Lee (2021) bibliography entry found: {text[:80]!r}")
                    break

        if lee_entry_found:
            total_score += 0.5
        else:
            print("FAIL: Component 1 — Lee (2021) bibliography entry not found in bibliography section")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # --- Component 2: Citation '(Lee, 2021)' inserted in 4th paragraph (0.5 points) ---
    # The task requires inserting a cross-reference/citation '(Lee, 2021)' in the 4th paragraph.
    # The 4th paragraph is counted as the 4th content paragraph (excluding headings),
    # which corresponds to doc.paragraphs[4] (index 4, after the Heading 1 at index 0).
    # The citation should appear somewhere in that paragraph.
    try:
        # Find the 4th content paragraph (skipping headings)
        content_para_count = 0
        fourth_para = None
        for para in doc.paragraphs:
            if not para.style.name.startswith('Heading'):
                content_para_count += 1
                if content_para_count == 4:
                    fourth_para = para
                    break

        if fourth_para is None:
            print("FAIL: Component 2 — Could not find 4th paragraph in document")
        else:
            para_text = fourth_para.text
            # Check for the citation pattern '(Lee, 2021)' or 'Lee, 2021' or 'Lee (2021)'
            import re
            citation_found = bool(
                re.search(r'\(Lee,?\s*2021\)', para_text) or
                re.search(r'Lee\s*\(2021\)', para_text)
            )
            if citation_found:
                # Find matching citation for display
                match = re.search(r'(\(Lee[^)]*2021[^)]*\)|Lee\s*\(2021\))', para_text)
                citation_str = match.group(0) if match else 'Lee...2021'
                print(f"PASS: Component 2 — Citation {citation_str!r} found in 4th paragraph (0.5 pts)")
                total_score += 0.5
            else:
                print(f"FAIL: Component 2 — Citation '(Lee, 2021)' not found in 4th paragraph")
                print(f"       4th paragraph text (last 100 chars): {para_text[-100:]!r}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in a given env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
