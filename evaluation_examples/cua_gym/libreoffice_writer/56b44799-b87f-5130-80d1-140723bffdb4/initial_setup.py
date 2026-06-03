"""
Initial Setup: Healthcare policy report with 5 paragraphs and 3 bibliography entries (no Lee 2021 entry).
Task ID: osworld_writer_bibliography_crossref_006
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_bibliography_crossref_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Healthcare Data Governance: Policy Frameworks and Implementation Challenges", level=1)

    # --- Paragraph 1 ---
    doc.add_paragraph(
        "The increasing digitization of healthcare records has created unprecedented opportunities for "
        "improving patient outcomes while simultaneously raising complex questions about data governance, "
        "privacy protections, and regulatory compliance. As health systems continue to adopt electronic "
        "health records (EHRs) and data-sharing platforms, policymakers must navigate the tension between "
        "enabling data-driven innovation and safeguarding patient privacy. Recent regulatory developments, "
        "including updated HIPAA guidelines and the European GDPR framework, have significantly reshaped "
        "institutional responsibilities regarding health data management (Smith, 2020)."
    )

    # --- Paragraph 2 ---
    doc.add_paragraph(
        "Interoperability remains one of the most pressing challenges in modern healthcare systems. "
        "The fragmentation of patient data across disparate systems, combined with inconsistent "
        "technical standards, hampers coordinated care delivery and population health management. "
        "Despite substantial federal investment through programs such as the 21st Century Cures Act, "
        "many healthcare organizations continue to struggle with siloed data architectures that impede "
        "timely clinical decision-making and longitudinal patient tracking (Johnson & Martinez, 2019). "
        "The adoption of Fast Healthcare Interoperability Resources (FHIR) standards represents a "
        "promising step toward resolving these challenges, though implementation timelines vary widely "
        "across institution types and geographic regions."
    )

    # --- Paragraph 3 ---
    doc.add_paragraph(
        "Artificial intelligence and machine learning applications in healthcare have generated "
        "significant research interest and clinical promise in recent years. Predictive analytics tools "
        "have demonstrated efficacy in early disease detection, treatment optimization, and hospital "
        "resource planning. However, the deployment of AI-driven clinical decision support systems "
        "raises important questions about algorithmic bias, explainability, and accountability. "
        "Healthcare institutions must develop robust governance frameworks that address data quality, "
        "model validation, and ongoing performance monitoring to ensure equitable patient outcomes "
        "across diverse demographic groups (Chen et al., 2022). Regulatory bodies are increasingly "
        "scrutinizing AI medical devices under existing frameworks, creating both compliance obligations "
        "and opportunities for thoughtful policy innovation."
    )

    # --- Paragraph 4 ---
    doc.add_paragraph(
        "The governance of secondary data use—encompassing research, quality improvement, and commercial "
        "applications—requires carefully designed consent frameworks and institutional oversight mechanisms. "
        "Patient trust is foundational to successful health data ecosystems; without meaningful consent "
        "processes and transparent data use disclosures, health systems risk eroding the public confidence "
        "necessary for sustaining broad participation in data-sharing initiatives. Dynamic consent models, "
        "which allow patients to granularly control data access and usage over time, offer a promising "
        "approach to balancing individual autonomy with population-level research imperatives. "
        "Institutional review boards and data governance committees must be adequately resourced and "
        "empowered to enforce compliance with evolving privacy standards."
    )

    # --- Paragraph 5 ---
    doc.add_paragraph(
        "Looking ahead, the convergence of genomic data, wearable device outputs, and social determinants "
        "of health information will require increasingly sophisticated governance architectures. "
        "Multi-stakeholder collaborations—spanning government agencies, health systems, technology vendors, "
        "and patient advocacy organizations—will be essential for developing governance standards that "
        "are both technically rigorous and socially accountable. Sustained investment in health informatics "
        "workforce development, combined with clear policy guidance on emerging technologies such as "
        "federated learning and differential privacy, will determine the extent to which healthcare "
        "systems can harness the full potential of data-driven medicine while preserving the ethical "
        "foundations upon which patient care depends."
    )

    # --- Bibliography Section ---
    bib_heading = doc.add_heading("Bibliography", level=2)

    doc.add_paragraph(
        "Chen, L., Patel, R., & Williams, D. (2022). Algorithmic fairness in clinical decision support: "
        "A systematic review. Journal of Biomedical Informatics, 128, 104034. "
        "https://doi.org/10.1016/j.jbi.2022.104034"
    )

    doc.add_paragraph(
        "Johnson, M., & Martinez, A. (2019). Health data interoperability: Barriers, progress, and "
        "policy recommendations. Health Affairs, 38(7), 1126–1134. "
        "https://doi.org/10.1377/hlthaff.2019.00512"
    )

    doc.add_paragraph(
        "Smith, T. (2020). Regulatory compliance and electronic health records: Navigating HIPAA "
        "in the age of big data. American Journal of Health Policy, 15(3), 89–105. "
        "https://doi.org/10.1093/ajhp/2020.15.3.89"
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
