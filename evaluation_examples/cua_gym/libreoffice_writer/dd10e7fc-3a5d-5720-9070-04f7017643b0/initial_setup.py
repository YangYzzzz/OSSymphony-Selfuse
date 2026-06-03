"""
Initial Setup: Set space before/after for Heading 2 paragraphs in operations manual
Task ID: writer_para_059
Domain: libreoffice_writer

Creates operations_manual.docx with Heading 2 paragraphs having NO spacing set
(space_before=0pt, space_after=0pt) — the agent must add 18pt before and 6pt after.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_para_059'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Title (Heading 1)
    heading1 = doc.add_heading('Standard Operating Procedures Manual', level=1)

    # Paragraph 2: Section heading (Heading 2) - space_before=0, space_after=0
    h2_receiving = doc.add_heading('Receiving and Inspection', level=2)
    h2_receiving.paragraph_format.space_before = Pt(0)
    h2_receiving.paragraph_format.space_after = Pt(0)

    # Paragraph 3: Body text
    doc.add_paragraph(
        'All incoming materials must be inspected within 24 hours of receipt. '
        'The receiving clerk shall verify the packing slip against the purchase order '
        'and note any discrepancies.'
    )

    # Paragraph 4: Body text
    doc.add_paragraph(
        'Items failing inspection must be quarantined in the designated rejection area '
        'and the supplier notified within 48 hours.'
    )

    # Paragraph 5: Section heading (Heading 2) - space_before=0, space_after=0
    h2_inventory = doc.add_heading('Inventory Management', level=2)
    h2_inventory.paragraph_format.space_before = Pt(0)
    h2_inventory.paragraph_format.space_after = Pt(0)

    # Paragraph 6: Body text
    doc.add_paragraph(
        'Physical inventory counts shall be conducted quarterly. '
        'Cycle counts of high-value items (Category A) shall be performed monthly.'
    )

    # Paragraph 7: Body text
    doc.add_paragraph(
        'All inventory adjustments exceeding $500 must be approved by the Operations Manager '
        'and documented with a written explanation.'
    )

    # Paragraph 8: Section heading (Heading 2) - space_before=0, space_after=0
    h2_shipping = doc.add_heading('Shipping and Fulfillment', level=2)
    h2_shipping.paragraph_format.space_before = Pt(0)
    h2_shipping.paragraph_format.space_after = Pt(0)

    # Paragraph 9: Body text
    doc.add_paragraph(
        'Orders must be picked, packed, and shipped within 24 hours of receipt for standard items. '
        'Custom orders follow the timeline specified in the customer agreement.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
