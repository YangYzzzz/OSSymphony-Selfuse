"""
Initial Setup: Apply consistent paragraph and heading styles across all subdocuments
Task ID: writer_rm_059
Domain: libreoffice_writer

Creates a master document and 4 subdocument chapters. Chapter2 has INCONSISTENT styles
(Calibri 16pt for Heading 1, Arial 11pt for Body Text) while all others use the master's
styles (Arial 18pt Bold for Heading 1, Times New Roman 12pt for Body Text).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_059'

# Master document styles
MASTER_HEADING_FONT = 'Arial'
MASTER_HEADING_SIZE = Pt(18)
MASTER_HEADING_BOLD = True
MASTER_BODY_FONT = 'Times New Roman'
MASTER_BODY_SIZE = Pt(12)

# Chapter2 inconsistent styles
CH2_HEADING_FONT = 'Calibri'
CH2_HEADING_SIZE = Pt(16)
CH2_HEADING_BOLD = False
CH2_BODY_FONT = 'Arial'
CH2_BODY_SIZE = Pt(11)


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, font_name, font_size, bold):
    """Add a heading paragraph with explicit font styling."""
    para = doc.add_paragraph()
    para.style = doc.styles['Heading 1']
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    return para


def add_body_styled(doc, text, font_name, font_size):
    """Add a body text paragraph with explicit font styling."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    return para


def create_master_document():
    """Create the master document that defines the style template."""
    doc = Document()
    master_path = f'{WORKDIR}/Corporate_Report_Master.docx'

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('Corporate Annual Report 2025')
    run.font.name = 'Arial'
    run.font.size = Pt(24)
    run.bold = True

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies Inc.')
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)

    doc.add_paragraph()  # spacer

    # Master style reference section
    add_heading_styled(doc, 'Document Style Guide', MASTER_HEADING_FONT, MASTER_HEADING_SIZE, MASTER_HEADING_BOLD)
    add_body_styled(doc, 'This master document defines the corporate style template for all subdocuments. '
                    'All chapters must use consistent heading and body text styles as defined below:', MASTER_BODY_FONT, MASTER_BODY_SIZE)
    add_body_styled(doc, 'Heading 1: Arial, 18pt, Bold', MASTER_BODY_FONT, MASTER_BODY_SIZE)
    add_body_styled(doc, 'Body Text: Times New Roman, 12pt', MASTER_BODY_FONT, MASTER_BODY_SIZE)

    doc.add_paragraph()

    # Subdocument listing
    add_heading_styled(doc, 'Subdocument Index', MASTER_HEADING_FONT, MASTER_HEADING_SIZE, MASTER_HEADING_BOLD)
    add_body_styled(doc, 'Chapter 1: Executive Summary', MASTER_BODY_FONT, MASTER_BODY_SIZE)
    add_body_styled(doc, 'Chapter 2: Financial Performance', MASTER_BODY_FONT, MASTER_BODY_SIZE)
    add_body_styled(doc, 'Chapter 3: Operations Review', MASTER_BODY_FONT, MASTER_BODY_SIZE)
    add_body_styled(doc, 'Chapter 4: Strategic Outlook', MASTER_BODY_FONT, MASTER_BODY_SIZE)

    doc.save(master_path)
    print(f'Master document created: {master_path}')


def create_chapter1():
    """Chapter 1 - Executive Summary (consistent styles)."""
    doc = Document()
    path = f'{WORKDIR}/Chapter1.docx'

    add_heading_styled(doc, 'Executive Summary', MASTER_HEADING_FONT, MASTER_HEADING_SIZE, MASTER_HEADING_BOLD)
    add_body_styled(doc, 'Meridian Technologies Inc. delivered another year of strong performance in fiscal year 2025. '
                    'Revenue grew by 18.3% year-over-year, reaching $4.72 billion, driven primarily by expansion '
                    'in our cloud services and enterprise software divisions.', MASTER_BODY_FONT, MASTER_BODY_SIZE)
    add_body_styled(doc, 'Operating income increased to $892 million, representing a 14.2% improvement from the prior year. '
                    'Our strategic investments in artificial intelligence and machine learning capabilities have begun '
                    'to yield significant returns, with the AI Solutions division contributing $340 million in revenue.', MASTER_BODY_FONT, MASTER_BODY_SIZE)

    add_heading_styled(doc, 'Key Highlights', MASTER_HEADING_FONT, MASTER_HEADING_SIZE, MASTER_HEADING_BOLD)
    add_body_styled(doc, 'Customer base expanded to 12,400 enterprise clients across 47 countries. '
                    'Employee headcount grew to 28,500, with notable hires in the research and development division. '
                    'Three new data centers were opened in Singapore, Frankfurt, and São Paulo to support global operations.', MASTER_BODY_FONT, MASTER_BODY_SIZE)
    add_body_styled(doc, 'The board of directors approved a quarterly dividend increase of 8%, reflecting confidence '
                    'in the company\'s sustained growth trajectory and strong cash flow generation.', MASTER_BODY_FONT, MASTER_BODY_SIZE)

    doc.save(path)
    print(f'Chapter 1 created: {path}')


def create_chapter2():
    """Chapter 2 - Financial Performance (INCONSISTENT styles - Calibri 16pt / Arial 11pt)."""
    doc = Document()
    path = f'{WORKDIR}/Chapter2.docx'

    # INCONSISTENT: Calibri 16pt (not bold) instead of Arial 18pt Bold
    add_heading_styled(doc, 'Financial Performance', CH2_HEADING_FONT, CH2_HEADING_SIZE, CH2_HEADING_BOLD)

    # INCONSISTENT: Arial 11pt instead of Times New Roman 12pt
    add_body_styled(doc, 'Total revenue for fiscal year 2025 reached $4.72 billion, an increase of 18.3% compared '
                    'to $3.99 billion in the prior year. This growth was driven by strong performance across all '
                    'major business segments, with cloud services leading at 24.7% growth.', CH2_BODY_FONT, CH2_BODY_SIZE)
    add_body_styled(doc, 'Gross profit margin expanded by 210 basis points to 62.4%, reflecting improved operational '
                    'efficiency and favorable product mix shifts toward higher-margin software subscriptions.', CH2_BODY_FONT, CH2_BODY_SIZE)

    add_heading_styled(doc, 'Revenue Breakdown by Segment', CH2_HEADING_FONT, CH2_HEADING_SIZE, CH2_HEADING_BOLD)
    add_body_styled(doc, 'Cloud Services: $1.89 billion (40.0% of total revenue, up from 37.9%)', CH2_BODY_FONT, CH2_BODY_SIZE)
    add_body_styled(doc, 'Enterprise Software: $1.42 billion (30.1%, stable year-over-year)', CH2_BODY_FONT, CH2_BODY_SIZE)
    add_body_styled(doc, 'Professional Services: $0.85 billion (18.0%, down from 19.4%)', CH2_BODY_FONT, CH2_BODY_SIZE)
    add_body_styled(doc, 'AI Solutions: $0.34 billion (7.2%, new segment)', CH2_BODY_FONT, CH2_BODY_SIZE)
    add_body_styled(doc, 'Other: $0.22 billion (4.7%)', CH2_BODY_FONT, CH2_BODY_SIZE)

    add_heading_styled(doc, 'Profitability Analysis', CH2_HEADING_FONT, CH2_HEADING_SIZE, CH2_HEADING_BOLD)
    add_body_styled(doc, 'Operating income reached $892 million, a 14.2% increase from the prior year. EBITDA totaled '
                    '$1.13 billion with a margin of 23.9%. Net income was $671 million, or $8.39 per diluted share, '
                    'compared to $587 million, or $7.34 per diluted share, in the prior year.', CH2_BODY_FONT, CH2_BODY_SIZE)
    add_body_styled(doc, 'Research and development expenses increased by 22.1% to $612 million, reflecting our '
                    'continued investment in next-generation AI and cloud infrastructure technologies.', CH2_BODY_FONT, CH2_BODY_SIZE)

    doc.save(path)
    print(f'Chapter 2 created (INCONSISTENT styles): {path}')


def create_chapter3():
    """Chapter 3 - Operations Review (consistent styles)."""
    doc = Document()
    path = f'{WORKDIR}/Chapter3.docx'

    add_heading_styled(doc, 'Operations Review', MASTER_HEADING_FONT, MASTER_HEADING_SIZE, MASTER_HEADING_BOLD)
    add_body_styled(doc, 'Our global operations infrastructure continued to scale efficiently throughout 2025. '
                    'System uptime across all production environments maintained a 99.97% availability rate, '
                    'exceeding our SLA commitment of 99.95%.', MASTER_BODY_FONT, MASTER_BODY_SIZE)
    add_body_styled(doc, 'The deployment of three new data centers in Singapore, Frankfurt, and São Paulo increased '
                    'our total compute capacity by 34%. Average response latency for cloud services improved by '
                    '18ms to 42ms globally, benefiting from edge computing optimizations.', MASTER_BODY_FONT, MASTER_BODY_SIZE)

    add_heading_styled(doc, 'Workforce Development', MASTER_HEADING_FONT, MASTER_HEADING_SIZE, MASTER_HEADING_BOLD)
    add_body_styled(doc, 'Total headcount reached 28,500 employees, with 3,200 net new hires during the fiscal year. '
                    'Engineering and R&D roles accounted for 58% of new positions. Employee retention rate '
                    'remained strong at 91.3%, supported by enhanced compensation packages and remote work flexibility.', MASTER_BODY_FONT, MASTER_BODY_SIZE)
    add_body_styled(doc, 'Our internal training programs delivered over 156,000 hours of professional development, '
                    'including specialized certifications in cloud architecture and AI engineering.', MASTER_BODY_FONT, MASTER_BODY_SIZE)

    add_heading_styled(doc, 'Supply Chain and Procurement', MASTER_HEADING_FONT, MASTER_HEADING_SIZE, MASTER_HEADING_BOLD)
    add_body_styled(doc, 'Strategic partnerships with key hardware vendors secured favorable pricing on server and '
                    'networking equipment, resulting in $47 million in procurement savings. Lead times for critical '
                    'infrastructure components were reduced by an average of 3.2 weeks through improved forecasting.', MASTER_BODY_FONT, MASTER_BODY_SIZE)

    doc.save(path)
    print(f'Chapter 3 created: {path}')


def create_chapter4():
    """Chapter 4 - Strategic Outlook (consistent styles)."""
    doc = Document()
    path = f'{WORKDIR}/Chapter4.docx'

    add_heading_styled(doc, 'Strategic Outlook', MASTER_HEADING_FONT, MASTER_HEADING_SIZE, MASTER_HEADING_BOLD)
    add_body_styled(doc, 'Looking ahead to fiscal year 2026, Meridian Technologies is well-positioned to capitalize '
                    'on accelerating demand for enterprise AI solutions and cloud infrastructure. We anticipate '
                    'revenue growth in the range of 15-20%, with AI Solutions expected to be the fastest-growing segment.', MASTER_BODY_FONT, MASTER_BODY_SIZE)
    add_body_styled(doc, 'Our three-year strategic plan, codenamed Project Horizon, outlines investments totaling '
                    '$2.1 billion in next-generation platform capabilities, including autonomous operations, '
                    'advanced analytics, and industry-specific vertical solutions.', MASTER_BODY_FONT, MASTER_BODY_SIZE)

    add_heading_styled(doc, 'Innovation Pipeline', MASTER_HEADING_FONT, MASTER_HEADING_SIZE, MASTER_HEADING_BOLD)
    add_body_styled(doc, 'Our research labs currently have 14 active projects in various stages of development, '
                    'spanning quantum-resistant cryptography, federated learning platforms, and real-time '
                    'natural language processing for enterprise workflows.', MASTER_BODY_FONT, MASTER_BODY_SIZE)
    add_body_styled(doc, 'Patent filings increased by 31% in 2025, with 127 new patents granted across AI, '
                    'cloud security, and data management technologies.', MASTER_BODY_FONT, MASTER_BODY_SIZE)

    add_heading_styled(doc, 'Market Expansion', MASTER_HEADING_FONT, MASTER_HEADING_SIZE, MASTER_HEADING_BOLD)
    add_body_styled(doc, 'Geographic expansion remains a priority, with planned market entries into South Korea, '
                    'the Middle East, and Sub-Saharan Africa during 2026. These regions represent a combined '
                    'addressable market of $18.5 billion for enterprise technology services.', MASTER_BODY_FONT, MASTER_BODY_SIZE)
    add_body_styled(doc, 'We are also exploring strategic acquisition opportunities to strengthen capabilities '
                    'in healthcare IT and financial services technology, two verticals with significant growth potential.', MASTER_BODY_FONT, MASTER_BODY_SIZE)

    doc.save(path)
    print(f'Chapter 4 created: {path}')


def create_initial():
    create_master_document()
    create_chapter1()
    create_chapter2()
    create_chapter3()
    create_chapter4()
    print('All files created successfully.')

    # Open the master document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{WORKDIR}/Corporate_Report_Master.docx"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
