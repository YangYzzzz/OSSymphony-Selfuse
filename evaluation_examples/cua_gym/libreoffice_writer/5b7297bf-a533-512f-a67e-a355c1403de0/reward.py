"""
Reward Script: Apply consistent paragraph and heading styles across all subdocuments
Task ID: writer_rm_059
Domain: libreoffice_writer
Scoring:
  Component 1 (0.40): Chapter2 Heading 1 paragraphs use Arial 18pt Bold
  Component 2 (0.40): Chapter2 Normal paragraphs use Times New Roman 12pt
  Component 3 (0.20): Style consistency — ALL runs in Chapter2 match expected styles
"""

import os
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_059'


def verify_task():
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: all required files must exist
    required_files = [
        'Corporate_Report_Master.docx',
        'Chapter1.docx',
        'Chapter2.docx',
        'Chapter3.docx',
        'Chapter4.docx',
    ]
    for f in required_files:
        fpath = os.path.join(WORKDIR, f)
        if not os.path.exists(fpath):
            print("CRITICAL: Required file missing: %s" % f)
            print("REWARD: 0.0")
            return 0.0

    # Load Chapter2 — the key file that should change
    ch2_path = os.path.join(WORKDIR, 'Chapter2.docx')
    try:
        doc = Document(ch2_path)
    except Exception as e:
        print("CRITICAL: Cannot load Chapter2.docx: %s" % e)
        print("REWARD: 0.0")
        return 0.0

    # Categorize paragraphs by style
    heading1_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Heading 1']
    normal_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Normal' and p.text.strip()]

    if not heading1_paras:
        print("FAIL: No Heading 1 paragraphs found in Chapter2.docx")
        print("REWARD: 0.0")
        return 0.0

    if not normal_paras:
        print("FAIL: No Normal paragraphs found in Chapter2.docx")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Chapter2 Heading 1 paragraphs use Arial 18pt Bold (0.40 points)
    # Initial state: Calibri 16pt, non-bold. Golden state: Arial 18pt Bold.
    try:
        h1_total = 0
        h1_pass = 0
        for para in heading1_paras:
            for run in para.runs:
                if not run.text.strip():
                    continue
                h1_total += 1
                font_ok = (run.font.name == 'Arial')
                size_ok = (run.font.size is not None and run.font.size.pt == 18.0)
                bold_ok = (run.font.bold is True)
                if font_ok and size_ok and bold_ok:
                    h1_pass += 1
                else:
                    print("FAIL detail: Heading 1 run '%s' — font=%s, size=%s, bold=%s" % (
                        run.text[:30], run.font.name,
                        run.font.size.pt if run.font.size else None,
                        run.font.bold))

        if h1_total > 0 and h1_pass == h1_total:
            print("PASS: Component 1 — All %d Heading 1 runs use Arial 18pt Bold (0.40 pts)" % h1_total)
            total_score += 0.40
        elif h1_total > 0:
            partial = 0.40 * (h1_pass / h1_total)
            print("PARTIAL: Component 1 — %d/%d Heading 1 runs correct (%.2f pts)" % (h1_pass, h1_total, partial))
            total_score += partial
        else:
            print("FAIL: Component 1 — No Heading 1 runs with text found")
    except Exception as e:
        print("ERROR: Component 1 — %s" % e)

    # Component 2: Chapter2 Normal paragraphs use Times New Roman 12pt (0.40 points)
    # Initial state: Arial 11pt. Golden state: Times New Roman 12pt.
    try:
        body_total = 0
        body_pass = 0
        for para in normal_paras:
            for run in para.runs:
                if not run.text.strip():
                    continue
                body_total += 1
                font_ok = (run.font.name == 'Times New Roman')
                size_ok = (run.font.size is not None and run.font.size.pt == 12.0)
                if font_ok and size_ok:
                    body_pass += 1
                else:
                    print("FAIL detail: Normal run '%s' — font=%s, size=%s" % (
                        run.text[:30], run.font.name,
                        run.font.size.pt if run.font.size else None))

        if body_total > 0 and body_pass == body_total:
            print("PASS: Component 2 — All %d Normal runs use Times New Roman 12pt (0.40 pts)" % body_total)
            total_score += 0.40
        elif body_total > 0:
            partial = 0.40 * (body_pass / body_total)
            print("PARTIAL: Component 2 — %d/%d Normal runs correct (%.2f pts)" % (body_pass, body_total, partial))
            total_score += partial
        else:
            print("FAIL: Component 2 — No Normal runs with text found")
    except Exception as e:
        print("ERROR: Component 2 — %s" % e)

    # Component 3: Full style consistency — no mixed/inconsistent styles remain (0.20 points)
    # This checks that EVERY run in Chapter2 matches the master template styles,
    # with zero deviations. Awards points only if both Component 1 and 2 fully passed.
    try:
        all_heading_correct = (h1_total > 0 and h1_pass == h1_total)
        all_body_correct = (body_total > 0 and body_pass == body_total)

        if all_heading_correct and all_body_correct:
            print("PASS: Component 3 — Full style consistency achieved (0.20 pts)")
            total_score += 0.20
        else:
            print("FAIL: Component 3 — Style inconsistencies remain (heading=%s, body=%s)" % (
                all_heading_correct, all_body_correct))
    except Exception as e:
        print("ERROR: Component 3 — %s" % e)

    final_score = round(min(total_score, 1.0), 2)
    print("")
    print("Score: %.2f/1.0" % total_score)
    print("REWARD: %.1f" % final_score)
    return final_score


# Entry point
verify_task()
