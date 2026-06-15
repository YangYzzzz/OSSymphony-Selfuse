"""
Initial Setup: Writer document with Quick Reference Card section in normal paragraph flow.
Task ID: writer_tech_075
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_075'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Linux System Administration Handbook', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Introduction ---
    doc.add_heading('Introduction', level=1)
    p = doc.add_paragraph(
        'This handbook provides a comprehensive overview of essential Linux system '
        'administration tasks. It covers user management, file system operations, '
        'network configuration, and security best practices for enterprise environments.'
    )

    p = doc.add_paragraph(
        'Whether you are a junior administrator setting up your first production server '
        'or a seasoned veteran managing a fleet of hundreds of machines, the techniques '
        'and procedures described here will serve as a reliable reference.'
    )

    # --- User Management ---
    doc.add_heading('User Management', level=1)
    doc.add_paragraph(
        'Managing user accounts is one of the most fundamental responsibilities of a '
        'system administrator. The following commands are used daily:'
    )
    doc.add_paragraph('useradd -m -s /bin/bash -G sudo newuser', style='List Bullet')
    doc.add_paragraph('passwd newuser', style='List Bullet')
    doc.add_paragraph('usermod -aG docker newuser', style='List Bullet')
    doc.add_paragraph('userdel -r olduser', style='List Bullet')

    doc.add_paragraph(
        'Always verify group memberships after changes using the "id" command. '
        'For bulk user provisioning, consider writing a script that reads from a CSV file.'
    )

    # --- File System Operations ---
    doc.add_heading('File System Operations', level=1)
    doc.add_paragraph(
        'Understanding disk layout and mount points is critical. Use lsblk and df -h to '
        'inspect current configurations. Common tasks include:'
    )

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Command', 'Purpose', 'Example']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['fdisk', 'Partition management', 'fdisk /dev/sdb'],
        ['mkfs.ext4', 'Create filesystem', 'mkfs.ext4 /dev/sdb1'],
        ['mount', 'Attach filesystem', 'mount /dev/sdb1 /mnt/data'],
        ['fstab', 'Persistent mounts', '/etc/fstab entry'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Network Configuration ---
    doc.add_heading('Network Configuration', level=1)
    doc.add_paragraph(
        'Modern Linux distributions use NetworkManager or systemd-networkd for network '
        'management. Key configuration files reside in /etc/NetworkManager/ or '
        '/etc/systemd/network/ depending on the distribution.'
    )
    doc.add_paragraph(
        'For static IP assignment on Ubuntu Server 22.04, edit the Netplan configuration '
        'at /etc/netplan/01-netcfg.yaml. Apply changes with "netplan apply" and verify '
        'with "ip addr show".'
    )

    # --- Quick Reference Card (normal paragraph flow, NO frame) ---
    doc.add_heading('Quick Reference Card', level=1)

    qr_intro = doc.add_paragraph(
        'The following commands are the most frequently used in daily operations:'
    )

    doc.add_paragraph('System Info: uname -a, hostnamectl, uptime', style='List Bullet')
    doc.add_paragraph('Process Mgmt: ps aux, top, htop, kill -9 PID', style='List Bullet')
    doc.add_paragraph('Disk Usage: df -h, du -sh /path, ncdu', style='List Bullet')
    doc.add_paragraph('Network: ip a, ss -tulnp, ping, traceroute', style='List Bullet')
    doc.add_paragraph('Logs: journalctl -xe, tail -f /var/log/syslog', style='List Bullet')
    doc.add_paragraph('Services: systemctl status|start|stop|restart <svc>', style='List Bullet')
    doc.add_paragraph('Package Mgmt: apt update && apt upgrade, dnf update', style='List Bullet')

    doc.add_paragraph(
        'Tip: Create shell aliases for your most-used commands in ~/.bashrc to save time.'
    )

    # --- Security Best Practices ---
    doc.add_heading('Security Best Practices', level=1)
    doc.add_paragraph(
        'Security is not an afterthought; it must be integrated into every stage of '
        'system setup and maintenance. Begin by disabling root SSH login and enforcing '
        'key-based authentication.'
    )
    doc.add_paragraph('Disable root login: PermitRootLogin no in /etc/ssh/sshd_config', style='List Bullet')
    doc.add_paragraph('Enable UFW: ufw enable && ufw allow 22/tcp', style='List Bullet')
    doc.add_paragraph('Automatic updates: apt install unattended-upgrades', style='List Bullet')

    doc.add_paragraph(
        'Regularly audit open ports with "ss -tulnp" and review authentication logs '
        'at /var/log/auth.log for suspicious activity.'
    )

    # --- Conclusion ---
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'System administration is a discipline that requires continuous learning. '
        'As infrastructure evolves toward containerization and cloud-native architectures, '
        'the fundamentals covered in this handbook remain essential building blocks.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
