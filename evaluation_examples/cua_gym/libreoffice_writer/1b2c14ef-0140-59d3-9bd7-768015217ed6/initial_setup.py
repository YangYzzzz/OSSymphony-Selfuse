"""
Initial Setup: Legal document with block quotes using 'Quotations' paragraph style.
Task ID: writer_fs_021
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_021'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # ---- Define 'Quotations' paragraph style ----
    styles = doc.styles
    quotations_style = styles.add_style('Quotations', 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    quotations_style.base_style = styles['Normal']
    fmt = quotations_style.paragraph_format
    fmt.left_indent = Inches(0.5)
    fmt.right_indent = Inches(0.5)
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # Deliberately NOT setting keep_together — that is the task
    qfont = quotations_style.font
    qfont.italic = True
    qfont.size = Pt(11)
    qfont.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ---- Page setup ----
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ---- Title ----
    title = doc.add_heading('Legal Analysis: Contract Enforceability Under the\nUniform Commercial Code', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')  # spacer

    # ---- Section 1 ----
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This memorandum examines the enforceability of commercial contracts under '
        'Article 2 of the Uniform Commercial Code (UCC), with particular focus on '
        'the parol evidence rule, the statute of frauds, and the implied duty of '
        'good faith and fair dealing. The analysis draws upon landmark case law and '
        'recent appellate decisions to provide a comprehensive framework for '
        'evaluating contract disputes in commercial transactions.'
    )
    doc.add_paragraph(
        'The significance of this area of law cannot be overstated, given the '
        'increasing complexity of modern business arrangements and the proliferation '
        'of electronic contracting mechanisms that challenge traditional doctrinal '
        'categories.'
    )

    # Block Quote 1
    doc.add_paragraph(
        '"The Uniform Commercial Code represents the most comprehensive attempt '
        'to rationalize and modernize the law governing commercial transactions. '
        'Its drafters sought to create a flexible framework that could accommodate '
        'the evolving practices of merchants while providing sufficient certainty '
        'to facilitate planning and reduce litigation costs." '
        '— Professor Eleanor Whitfield, Commercial Law Treatise (4th ed. 2019), at 47.',
        style='Quotations'
    )

    doc.add_paragraph(
        'This principle has been consistently affirmed by courts across multiple '
        'jurisdictions, reflecting a broad consensus regarding the purposes '
        'underlying the statutory scheme.'
    )

    # ---- Section 2 ----
    doc.add_heading('2. The Parol Evidence Rule Under UCC Section 2-202', level=1)
    doc.add_paragraph(
        'UCC Section 2-202 modifies the common law parol evidence rule by permitting '
        'the introduction of course of dealing, usage of trade, and course of performance '
        'to supplement or explain the terms of a written agreement, even when that '
        'agreement appears on its face to be complete and unambiguous. This departure '
        'from classical contract doctrine has generated substantial litigation, '
        'particularly in cases involving sophisticated commercial parties.'
    )

    # Block Quote 2
    doc.add_paragraph(
        '"Terms with respect to which the confirmatory memoranda of the parties agree '
        'or which are otherwise set forth in a writing intended by the parties as a '
        'final expression of their agreement with respect to such terms as are included '
        'therein may not be contradicted by evidence of any prior agreement or of a '
        'contemporaneous oral agreement but may be explained or supplemented by course '
        'of dealing or usage of trade or by course of performance." '
        '— UCC Section 2-202.',
        style='Quotations'
    )

    doc.add_paragraph(
        'The practical import of this provision is that parties to a commercial contract '
        'cannot rely solely on a merger clause to exclude extrinsic evidence of trade '
        'usage or course of dealing. This has significant implications for contract '
        'drafting and dispute resolution strategies.'
    )

    doc.add_paragraph(
        'Courts have taken divergent approaches to the question of when a writing '
        'should be deemed a complete and exclusive statement of the parties\' agreement. '
        'The majority view, articulated most clearly by the Second Circuit, holds that '
        'the determination must be made in light of all surrounding circumstances.'
    )

    # Block Quote 3
    doc.add_paragraph(
        '"Where the parties have reduced their agreement to writing, the question of '
        'whether that writing was intended as the complete and exclusive statement of '
        'the terms of their contract is one of fact, to be determined from the writing '
        'itself, the testimony of the parties, and the surrounding circumstances at the '
        'time of contracting, including applicable trade usage and prior dealings between '
        'the parties." '
        '— Columbia Nitrogen Corp. v. Royster Co., 451 F.2d 3, 9 (4th Cir. 1971).',
        style='Quotations'
    )

    # ---- Section 3 ----
    doc.add_heading('3. Statute of Frauds: UCC Section 2-201', level=1)
    doc.add_paragraph(
        'Section 2-201 of the UCC imposes a writing requirement for contracts '
        'involving the sale of goods valued at $500 or more. Unlike the common law '
        'statute of frauds, however, the UCC version is considerably more lenient, '
        'requiring only a writing "sufficient to indicate that a contract for sale '
        'has been made between the parties." The writing need not contain all material '
        'terms; indeed, the only term that must appear is the quantity.'
    )
    doc.add_paragraph(
        'Three significant exceptions to the writing requirement deserve particular '
        'attention in the context of modern commercial practice.'
    )

    # Block Quote 4
    doc.add_paragraph(
        '"Between merchants, if within a reasonable time a writing in confirmation of '
        'the contract and sufficient against the sender is received and the party '
        'receiving it has reason to know its contents, it satisfies the requirements '
        'of subsection (1) against such party unless written notice of objection to '
        'its contents is given within 10 days after it is received." '
        '— UCC Section 2-201(2), the merchant\'s exception.',
        style='Quotations'
    )

    doc.add_paragraph(
        'This provision has particular relevance in the era of electronic communications, '
        'where confirmatory emails and electronic purchase orders serve the same function '
        'as traditional written confirmations. Courts have uniformly held that electronic '
        'communications satisfy the merchant\'s exception, provided the basic requirements '
        'of the statute are otherwise met.'
    )

    # ---- Section 4 ----
    doc.add_heading('4. Good Faith and Fair Dealing', level=1)
    doc.add_paragraph(
        'UCC Section 1-304 imposes an obligation of good faith in the performance '
        'and enforcement of every contract or duty within the scope of the Code. For '
        'merchants, good faith is defined to include "honesty in fact and the observance '
        'of reasonable commercial standards of fair dealing in the trade." This heightened '
        'standard reflects the commercial law\'s emphasis on protecting the reasonable '
        'expectations of parties in ongoing business relationships.'
    )

    # Block Quote 5
    doc.add_paragraph(
        '"The obligation of good faith does not serve as a roving commission to the '
        'courts to rewrite contracts in the interest of fairness. Rather, it functions '
        'as a gap-filler, supplying terms that the parties would have agreed upon had '
        'they anticipated the situation that subsequently arose. The doctrine prevents '
        'one party from exploiting the literal terms of the contract to deprive the '
        'other of the benefit of the bargain." '
        '— Market Street Associates Ltd. v. Frey, 941 F.2d 588, 595 (7th Cir. 1991) (Posner, J.).',
        style='Quotations'
    )

    doc.add_paragraph(
        'Judge Posner\'s formulation has been widely cited and adopted, establishing '
        'a framework that balances contractual freedom with the need to prevent '
        'opportunistic behavior. The practical significance of this doctrine is perhaps '
        'most apparent in requirements and output contracts, where the buyer or seller '
        'retains substantial discretion in determining the quantity of goods to be '
        'purchased or supplied.'
    )

    # Block Quote 6
    doc.add_paragraph(
        '"A requirements contract does not give the buyer unlimited discretion to '
        'determine the quantity to be purchased. The buyer must exercise good faith in '
        'setting its requirements, and any quantity unreasonably disproportionate to a '
        'stated estimate or to any normal or otherwise comparable prior requirements '
        'is not enforceable. The standard is one of commercial reasonableness, measured '
        'against the legitimate business reasons underlying the buyer\'s decision." '
        '— Feld v. Henry S. Levy & Sons, Inc., 37 N.Y.2d 466, 471 (1975).',
        style='Quotations'
    )

    # ---- Section 5 ----
    doc.add_heading('5. Remedial Framework', level=1)
    doc.add_paragraph(
        'The UCC provides a comprehensive remedial framework that seeks to place the '
        'aggrieved party in as good a position as performance would have. Article 2 '
        'distinguishes between buyer\'s remedies (Sections 2-711 through 2-717) and '
        'seller\'s remedies (Sections 2-703 through 2-710), reflecting the different '
        'interests at stake depending on which party has breached.'
    )
    doc.add_paragraph(
        'The centerpiece of the buyer\'s remedial scheme is the cover provision, '
        'which permits the buyer to procure substitute goods and recover the difference '
        'between the cover price and the contract price.'
    )

    # Block Quote 7
    doc.add_paragraph(
        '"After a breach, the buyer may \'cover\' by making in good faith and without '
        'unreasonable delay any reasonable purchase of or contract for goods in '
        'substitution for those due from the seller. The buyer may recover from the '
        'seller as damages the difference between the cost of cover and the contract '
        'price together with any incidental and consequential damages, but less expenses '
        'saved in consequence of the seller\'s breach." '
        '— UCC Section 2-712.',
        style='Quotations'
    )

    doc.add_paragraph(
        'The cover remedy has been praised for its practical orientation and its '
        'alignment with commercial reality. Unlike the traditional contract-market '
        'differential measure, cover focuses on the actual economic position of the '
        'aggrieved buyer, providing a more accurate assessment of compensable loss.'
    )

    # ---- Section 6 ----
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        'The Uniform Commercial Code\'s treatment of contract enforceability reflects '
        'a carefully calibrated balance between formality and flexibility, certainty and '
        'fairness. As commercial transactions continue to evolve in response to '
        'technological innovation and changing market practices, the foundational '
        'principles discussed in this memorandum will remain essential to the resolution '
        'of contractual disputes.'
    )

    # Block Quote 8
    doc.add_paragraph(
        '"The great achievement of the Code\'s contract provisions lies in their capacity '
        'to accommodate the infinite variety of commercial arrangements that human '
        'ingenuity can devise, while providing a stable framework of default rules that '
        'reduce transaction costs and promote efficient exchange. The Code is, at its '
        'core, an exercise in practical jurisprudence." '
        '— Professor Arthur Leff, Yale Law Journal (1967), at 185.',
        style='Quotations'
    )

    doc.add_paragraph(
        'Future developments in artificial intelligence, blockchain-based smart contracts, '
        'and decentralized autonomous organizations will undoubtedly test the limits of '
        'existing doctrinal categories. The flexibility inherent in the UCC\'s approach, '
        'however, suggests that it will continue to provide a workable framework for '
        'resolving the disputes that inevitably arise from the complex web of modern '
        'commercial relationships.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
