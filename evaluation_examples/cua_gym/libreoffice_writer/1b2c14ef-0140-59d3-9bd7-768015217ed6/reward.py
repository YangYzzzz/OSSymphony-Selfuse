"""
Reward Script: Set 'Do not split paragraph' for all block quotes in legal document
Task ID: writer_fs_021
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Quotations style definition has keep_together (keepLines) enabled
  Component 2 (0.3): All Quotations paragraphs effectively have keep_together
                      (either via style inheritance or direct paragraph formatting)
  Component 3 (0.2): Document integrity - still has 8 Quotations paragraphs with content
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_021'
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def check_keep_lines_on_style(style_element):
    """Check if a style element has keepLines enabled in its pPr."""
    ppr = style_element.find('.//w:pPr', NS)
    if ppr is None:
        return False
    kl = ppr.find('w:keepLines', NS)
    if kl is None:
        return False
    # In OOXML, presence of <w:keepLines/> means True.
    # <w:keepLines w:val="0"/> or w:val="false" means disabled.
    val = kl.get(qn('w:val'))
    if val is not None and val.lower() in ('0', 'false'):
        return False
    return True


def check_keep_lines_on_paragraph(para_element):
    """Check if a paragraph element has keepLines directly set in its pPr."""
    ppr = para_element.find('w:pPr', NS)
    if ppr is None:
        return False
    kl = ppr.find('w:keepLines', NS)
    if kl is None:
        return False
    val = kl.get(qn('w:val'))
    if val is not None and val.lower() in ('0', 'false'):
        return False
    return True


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Quotations style has keep_together/keepLines enabled (0.5 points)
    # This is the primary task requirement - modifying the style definition.
    try:
        style = doc.styles['Quotations']
        style_has_keep_lines = check_keep_lines_on_style(style.element)

        if style_has_keep_lines:
            print(f"PASS: Component 1 -- Quotations style has keepLines enabled (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 -- Quotations style does NOT have keepLines enabled")
    except KeyError:
        print(f"ERROR: Component 1 -- 'Quotations' style not found in document")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: All Quotations paragraphs effectively have keep_together (0.3 points)
    # The task says ALL block quotes should not be split. Check that each paragraph
    # with Quotations style has keepLines either via style or direct formatting.
    # We check both: style-level (already checked above) OR paragraph-level override.
    try:
        quotation_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Quotations']
        if len(quotation_paras) == 0:
            print(f"FAIL: Component 2 -- No paragraphs with Quotations style found")
        else:
            fail_count = 0
            for i, para in enumerate(quotation_paras):
                # Effective keep_together: either style defines it, or paragraph overrides it
                para_direct = check_keep_lines_on_paragraph(para._element)
                # python-docx keep_together resolves inheritance
                pf_value = para.paragraph_format.keep_together

                # Effective if: style has it (Component 1), OR paragraph directly has it,
                # OR python-docx resolves it as True
                effective = style_has_keep_lines or para_direct or (pf_value is True)

                if not effective:
                    fail_count += 1
                    print(f"  Para '{para.text[:40]}...' does NOT have effective keepLines")

            if fail_count == 0:
                print(f"PASS: Component 2 -- All {len(quotation_paras)} Quotations paragraphs have effective keepLines (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 -- Not all Quotations paragraphs have effective keepLines")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Document integrity - still has 8 Quotations paragraphs (0.2 points)
    # Ensures the agent didn't delete or change the style of block quotes.
    # This ONLY scores if keepLines is set (anchored to task change).
    try:
        quotation_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Quotations']
        num_quotes = len(quotation_paras)
        has_content = all(len(p.text.strip()) > 0 for p in quotation_paras)

        # Only award if keepLines is actually set (task change happened)
        if style_has_keep_lines and num_quotes == 8 and has_content:
            print(f"PASS: Component 3 -- Document integrity: {num_quotes} Quotations paragraphs with content preserved (0.2 pts)")
            total_score += 0.2
        elif not style_has_keep_lines:
            print(f"FAIL: Component 3 -- keepLines not set, so integrity check not awarded")
        elif num_quotes != 8:
            print(f"FAIL: Component 3 -- Expected 8 Quotations paragraphs, found {num_quotes}")
        else:
            print(f"FAIL: Component 3 -- Some Quotations paragraphs have empty content")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
