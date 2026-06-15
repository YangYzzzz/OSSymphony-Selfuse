"""
Initial Setup: Literature review document with target sentence for footnote insertion
Task ID: writer_acad_006
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    title = doc.add_heading('Literature Review: Organizational Behavior and Employee Motivation', level=1)

    # Paragraph 1: Introduction
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(12)
    p1.paragraph_format.line_spacing = 1.5
    run1 = p1.add_run(
        'The study of organizational behavior has evolved significantly over the past several decades. '
        'Researchers have explored various dimensions of workplace dynamics, from leadership styles to '
        'team cohesion and individual performance metrics. Early investigations focused primarily on '
        'structural determinants of productivity, while more recent scholarship has emphasized the '
        'psychological and social factors that influence employee engagement and satisfaction. '
        'This literature review synthesizes key findings from the field, tracing the development '
        'of major theoretical frameworks and their empirical support.'
    )
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)

    # Paragraph 2: Contains the target sentence
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    p2.paragraph_format.line_spacing = 1.5
    run2 = p2.add_run(
        'Among the most influential contributions to motivation research is the expectancy-valence model, '
        'which posits that employee effort is a function of perceived outcomes and their desirability. '
        'This theory was first proposed in 1987. '
        'Subsequent studies have refined the model by incorporating moderating variables such as '
        'self-efficacy, organizational culture, and reward structure. The expectancy-valence framework '
        'remains a cornerstone of contemporary motivation research, informing both academic inquiry '
        'and practical interventions in human resource management.'
    )
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    # Paragraph 3: Additional content
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(12)
    p3.paragraph_format.line_spacing = 1.5
    run3 = p3.add_run(
        'A parallel line of inquiry has examined the role of transformational leadership in fostering '
        'intrinsic motivation. Burns (1978) originally distinguished between transactional and '
        'transformational leadership styles, arguing that the latter generates deeper commitment '
        'and higher performance among subordinates. Bass and Avolio (1994) operationalized this '
        'distinction through the Multifactor Leadership Questionnaire, enabling large-scale empirical '
        'testing across industries and cultural contexts. Meta-analytic reviews have generally supported '
        'the positive association between transformational leadership and a range of desirable outcomes, '
        'including job satisfaction, organizational citizenship behavior, and reduced turnover intention.'
    )
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)

    # Paragraph 4: More content
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(12)
    p4.paragraph_format.line_spacing = 1.5
    run4 = p4.add_run(
        'More recently, scholars have turned their attention to the intersection of technology and '
        'organizational behavior. The rapid adoption of remote work technologies during the 2020 pandemic '
        'created natural experiments that allowed researchers to study the effects of physical separation '
        'on team dynamics, communication patterns, and individual well-being. Preliminary findings '
        'suggest that while remote work can enhance flexibility and autonomy, it may also erode '
        'informal social networks and reduce opportunities for spontaneous collaboration. These tensions '
        'highlight the need for adaptive management strategies that balance efficiency with employee '
        'connection and belonging.'
    )
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(12)

    # Paragraph 5: Conclusion
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(12)
    p5.paragraph_format.line_spacing = 1.5
    run5 = p5.add_run(
        'In summary, the field of organizational behavior continues to expand in both scope and '
        'methodological sophistication. Future research should aim to integrate insights from '
        'cognitive psychology, data science, and cross-cultural studies to build more comprehensive '
        'models of workplace behavior. The practical implications of this body of work are substantial, '
        'as organizations worldwide seek evidence-based approaches to talent management, employee '
        'development, and organizational design.'
    )
    run5.font.name = 'Times New Roman'
    run5.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
