"""
Initial Setup: Academic paper with Escherichia coli occurrences in regular text.
Task ID: writer_acad_065
Domain: libreoffice_writer

Creates a Writer document with 8 occurrences of 'Escherichia coli' in regular
(non-italic) text. A character style 'Latin Name' with italic exists but is
NOT applied to any text.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_065'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_character_style_latin_name(doc):
    """
    Add a character style 'Latin Name' with italic=True to the document styles.
    python-docx does not support character styles natively, so we manipulate the XML.
    """
    styles_element = doc.styles.element

    # Create a new character style element
    style_xml = (
        '<w:style w:type="character" w:customStyle="1" w:styleId="LatinName" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '  <w:name w:val="Latin Name"/>'
        '  <w:rPr>'
        '    <w:i/>'
        '    <w:iCs/>'
        '  </w:rPr>'
        '</w:style>'
    )
    style_elem = parse_xml(style_xml)
    styles_element.append(style_elem)


def add_paragraph_with_style(doc, text, style_name=None, bold=False, font_size=None, alignment=None):
    """Helper to add a styled paragraph."""
    para = doc.add_paragraph()
    if style_name:
        para.style = doc.styles[style_name]
    if alignment is not None:
        para.paragraph_format.alignment = alignment
    run = para.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    return para


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Add the 'Latin Name' character style (italic)
    add_character_style_latin_name(doc)

    # --- Title ---
    title = doc.add_heading('', level=0)
    title_run = title.add_run(
        'Antibiotic Resistance Mechanisms in Escherichia coli: '
        'A Comprehensive Review of Multidrug Efflux Systems'
    )
    title_run.font.size = Pt(16)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Authors ---
    authors = doc.add_paragraph()
    authors.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    authors.paragraph_format.space_after = Pt(6)
    run = authors.add_run(
        'Dr. Amara Okafor, Dr. Rajesh Patel, Dr. Mei-Lin Zhang\n'
        'Department of Microbiology, University of Cambridge\n'
        'Correspondence: a.okafor@cam.ac.uk'
    )
    run.font.size = Pt(10)

    # --- Abstract ---
    # Occurrence 1 in title already, let's count carefully:
    # Title: 1 occurrence
    # Abstract: 2 occurrences (total 3)
    # Introduction: 2 occurrences (total 5)
    # Methods: 1 occurrence (total 6)
    # Results: 1 occurrence (total 7)
    # Discussion: 1 occurrence (total 8)

    abstract_heading = doc.add_heading('Abstract', level=1)

    abstract_para = doc.add_paragraph()
    abstract_para.paragraph_format.space_after = Pt(6)
    run = abstract_para.add_run(
        'The increasing prevalence of antibiotic-resistant bacterial strains poses a '
        'significant threat to global public health. Among Gram-negative pathogens, '
        'Escherichia coli remains one of the most clinically relevant organisms due to '
        'its role in urinary tract infections, septicemia, and neonatal meningitis. '
        'This review examines the molecular mechanisms underlying multidrug resistance '
        'in Escherichia coli, with particular emphasis on the AcrAB-TolC efflux pump system. '
        'We analyze recent crystallographic data, genetic studies, and clinical isolate '
        'characterizations to provide a comprehensive overview of current knowledge '
        'in this critical area of antimicrobial research.'
    )

    # --- Keywords ---
    kw_para = doc.add_paragraph()
    kw_run = kw_para.add_run('Keywords: ')
    kw_run.bold = True
    kw_para.add_run(
        'antibiotic resistance, efflux pumps, AcrAB-TolC, multidrug resistance, '
        'Gram-negative bacteria, antimicrobial susceptibility'
    )

    # --- 1. Introduction ---
    doc.add_heading('1. Introduction', level=1)

    intro1 = doc.add_paragraph()
    intro1.add_run(
        'Antimicrobial resistance (AMR) has been identified by the World Health Organization '
        'as one of the top ten global public health threats facing humanity. The emergence '
        'and spread of resistant bacterial strains have outpaced the development of novel '
        'antimicrobial agents, creating an urgent need for deeper understanding of resistance '
        'mechanisms. Escherichia coli, a Gram-negative facultative anaerobe belonging to the '
        'family Enterobacteriaceae, serves as both a commensal organism in the human gut '
        'and a versatile pathogen capable of causing diverse infections.'
    )

    intro2 = doc.add_paragraph()
    intro2.add_run(
        'Pathogenic strains of Escherichia coli have acquired multiple mechanisms of '
        'antibiotic resistance, including enzymatic inactivation of antimicrobial compounds, '
        'modification of drug targets, decreased membrane permeability, and active drug efflux. '
        'Of these, multidrug efflux systems represent a particularly concerning mechanism due '
        'to their ability to confer simultaneous resistance to multiple structurally unrelated '
        'antimicrobial agents. The AcrAB-TolC system, a tripartite resistance-nodulation-division '
        '(RND) family efflux pump, has emerged as the primary contributor to intrinsic and '
        'acquired multidrug resistance in this organism.'
    )

    # --- 2. Materials and Methods ---
    doc.add_heading('2. Materials and Methods', level=1)

    methods1 = doc.add_paragraph()
    methods1.add_run(
        'Clinical isolates of Escherichia coli were obtained from the University Hospital '
        'microbiology laboratory between January 2023 and December 2024. A total of 342 '
        'non-duplicate isolates were collected from blood cultures (n=89), urine samples '
        '(n=178), wound swabs (n=45), and cerebrospinal fluid (n=30). Species identification '
        'was confirmed using MALDI-TOF mass spectrometry (Bruker Daltonics, Bremen, Germany). '
        'Antimicrobial susceptibility testing was performed according to EUCAST guidelines '
        'using the broth microdilution method.'
    )

    methods2 = doc.add_paragraph()
    methods2.add_run(
        'Genomic DNA was extracted using the QIAamp DNA Mini Kit (Qiagen, Hilden, Germany). '
        'Whole-genome sequencing was performed on an Illumina NovaSeq 6000 platform with '
        '150 bp paired-end reads. Resistance gene identification was carried out using '
        'ResFinder 4.1 and the Comprehensive Antibiotic Resistance Database (CARD). '
        'Efflux pump gene expression levels were quantified by RT-qPCR using species-specific '
        'primers designed against the acrA, acrB, and tolC genes.'
    )

    # --- 3. Results ---
    doc.add_heading('3. Results', level=1)

    # Add a simple table for results
    results_intro = doc.add_paragraph()
    results_intro.add_run(
        'Of the 342 clinical isolates analyzed, 187 (54.7%) demonstrated multidrug resistance, '
        'defined as non-susceptibility to at least one agent in three or more antimicrobial '
        'categories. Table 1 summarizes the resistance profiles observed among the '
        'Escherichia coli isolates.'
    )

    # Table 1
    table_caption = doc.add_paragraph()
    tc_run = table_caption.add_run('Table 1. ')
    tc_run.bold = True
    table_caption.add_run(
        'Antimicrobial resistance rates among clinical isolates (n=342).'
    )

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['Antimicrobial Agent', 'Resistant (n)', 'Resistance Rate (%)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    data = [
        ['Ampicillin', '267', '78.1'],
        ['Ciprofloxacin', '156', '45.6'],
        ['Ceftriaxone', '98', '28.7'],
        ['Trimethoprim-sulfamethoxazole', '203', '59.4'],
        ['Gentamicin', '84', '24.6'],
        ['Meropenem', '12', '3.5'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    results2 = doc.add_paragraph()
    results2.paragraph_format.space_before = Pt(12)
    results2.add_run(
        'RT-qPCR analysis revealed that acrB expression was upregulated 4.2-fold (SD +/- 1.3) '
        'in multidrug-resistant isolates compared to susceptible controls (p < 0.001). '
        'Furthermore, 23 isolates harboring mutations in the acrR regulatory gene showed '
        'significantly elevated minimum inhibitory concentrations (MICs) for ciprofloxacin, '
        'tetracycline, and chloramphenicol.'
    )

    # --- 4. Discussion ---
    doc.add_heading('4. Discussion', level=1)

    disc1 = doc.add_paragraph()
    disc1.add_run(
        'The findings of this study underscore the critical role of efflux-mediated resistance '
        'in clinical Escherichia coli isolates. The high prevalence of multidrug resistance '
        'observed in our cohort (54.7%) aligns with recent reports from surveillance programs '
        'across Europe and Asia. The strong correlation between acrB overexpression and '
        'elevated MICs confirms the AcrAB-TolC system as a primary determinant of clinically '
        'relevant resistance in this species.'
    )

    disc2 = doc.add_paragraph()
    disc2.add_run(
        'Our crystallographic analysis of the AcrB transporter protein revealed novel insights '
        'into substrate binding and the peristaltic mechanism of drug extrusion. The identification '
        'of a previously uncharacterized binding pocket in the periplasmic domain suggests '
        'potential targets for efflux pump inhibitor (EPI) development. Combination therapy '
        'incorporating EPIs with existing antibiotics may represent a viable strategy to '
        'restore susceptibility in resistant strains.'
    )

    # --- References ---
    doc.add_heading('References', level=1)
    refs = [
        'Blair, J.M.A., Richmond, G.E., & Piddock, L.J.V. (2014). Multidrug efflux pumps in Gram-negative bacteria and their role in antibiotic resistance. Future Microbiology, 9(10), 1165-1177.',
        'Du, D., Wang-Kan, X., Neuberger, A., et al. (2018). Multidrug efflux pumps: structure, function and regulation. Nature Reviews Microbiology, 16(9), 523-539.',
        'Poole, K. (2005). Efflux-mediated antimicrobial resistance. Journal of Antimicrobial Chemotherapy, 56(1), 20-51.',
        'Nikaido, H., & Pages, J.M. (2012). Broad-specificity efflux pumps and their role in multidrug resistance of Gram-negative bacteria. FEMS Microbiology Reviews, 36(2), 340-363.',
        'Webber, M.A., & Piddock, L.J.V. (2003). The importance of efflux pumps in bacterial antibiotic resistance. Journal of Antimicrobial Chemotherapy, 51(1), 9-11.',
    ]
    for i, ref in enumerate(refs, 1):
        ref_para = doc.add_paragraph()
        ref_para.add_run(f'[{i}] {ref}')
        ref_para.paragraph_format.space_after = Pt(3)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Count occurrences to verify
    full_text = []
    for p in doc.paragraphs:
        full_text.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text.append(p.text)
    all_text = '\n'.join(full_text)
    count = all_text.count('Escherichia coli')
    print(f'Occurrences of "Escherichia coli": {count}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
