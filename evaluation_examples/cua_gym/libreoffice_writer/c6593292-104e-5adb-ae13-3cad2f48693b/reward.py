"""
Reward Script: Apply 'Latin Name' character style to all 'Escherichia coli' occurrences
Task ID: writer_acad_065
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): All 8 occurrences are italic
  Component 2 (0.4): All 8 occurrences have 'Latin Name' character style
  Component 3 (0.2): Text content is preserved (no data corruption)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_065'
EXPECTED_OCCURRENCES = 8
TARGET_TEXT = 'Escherichia coli'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def find_ecoli_runs(doc):
    """
    Find all runs that contain 'Escherichia coli' text.
    Returns list of (paragraph_index, run) tuples.

    Handles two cases:
    1. 'Escherichia coli' is in a single run
    2. 'Escherichia coli' is split across the paragraph text but we find it
       by scanning run text.
    """
    results = []
    for pi, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if TARGET_TEXT in run.text:
                results.append((pi, run))
    return results


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find all runs containing 'Escherichia coli'
    ecoli_runs = find_ecoli_runs(doc)
    found_count = len(ecoli_runs)
    print(f"INFO: Found {found_count} runs containing '{TARGET_TEXT}'")

    # Also count occurrences in full text to cross-check
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    text_occurrences = full_text.count(TARGET_TEXT)
    print(f"INFO: Full text contains {text_occurrences} occurrences of '{TARGET_TEXT}'")

    if text_occurrences < EXPECTED_OCCURRENCES:
        print(f"FAIL: Expected at least {EXPECTED_OCCURRENCES} occurrences in text, found {text_occurrences}. Possible data corruption.")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All occurrences are italic (0.4 points)
    # This checks that each run containing 'Escherichia coli' has italic=True
    try:
        if found_count == 0:
            print(f"FAIL: Component 1 — No dedicated runs found for '{TARGET_TEXT}' (text not separated into styled runs)")
            italic_fraction = 0.0
        else:
            italic_count = sum(1 for _, run in ecoli_runs if run.font.italic is True)
            italic_fraction = italic_count / found_count
            if italic_fraction == 1.0:
                print(f"PASS: Component 1 — All {found_count} runs are italic (0.4 pts)")
                total_score += 0.4
            elif italic_fraction > 0:
                partial = round(0.4 * italic_fraction, 2)
                print(f"PARTIAL: Component 1 — {italic_count}/{found_count} runs are italic ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 1 — 0/{found_count} runs are italic")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All occurrences have 'Latin Name' character style (0.4 points)
    # This is the core task requirement: applying the specific character style
    try:
        if found_count == 0:
            print(f"FAIL: Component 2 — No dedicated runs found for '{TARGET_TEXT}'")
            style_fraction = 0.0
        else:
            style_count = 0
            for _, run in ecoli_runs:
                style_name = run.style.name if run.style else None
                if style_name == 'Latin Name':
                    style_count += 1
            style_fraction = style_count / found_count
            if style_fraction == 1.0:
                print(f"PASS: Component 2 — All {found_count} runs have 'Latin Name' style (0.4 pts)")
                total_score += 0.4
            elif style_fraction > 0:
                partial = round(0.4 * style_fraction, 2)
                print(f"PARTIAL: Component 2 — {style_count}/{found_count} runs have 'Latin Name' style ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — 0/{found_count} runs have 'Latin Name' style")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Runs are properly isolated — 'Escherichia coli' exists as a
    # dedicated run (not lumped with surrounding text) AND italic+styled (0.2 points)
    # In initial_env, the text is part of larger runs, so dedicated run count == 0.
    # In golden_env, each occurrence should be in its own clean run.
    try:
        dedicated_count = 0
        for _, run in ecoli_runs:
            # A "dedicated" run means the run text is exactly 'Escherichia coli'
            # or at least the run has the style AND italic applied.
            # In initial, the run contains entire paragraph text, so this won't match.
            run_style = run.style.name if run.style else None
            if run.font.italic is True and run_style == 'Latin Name':
                dedicated_count += 1
        if dedicated_count >= EXPECTED_OCCURRENCES:
            print(f"PASS: Component 3 — {dedicated_count} properly styled+italic dedicated runs (0.2 pts)")
            total_score += 0.2
        elif dedicated_count > 0:
            partial = round(0.2 * (dedicated_count / EXPECTED_OCCURRENCES), 2)
            print(f"PARTIAL: Component 3 — {dedicated_count}/{EXPECTED_OCCURRENCES} properly styled+italic runs ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No properly styled+italic dedicated runs found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
