"""
Reward Script: Mail merge DL envelopes for 25 vendors with company logo
Task ID: writer_mt_013
Domain: libreoffice_writer
Scoring:
  Component 1: DL envelope dimensions (220mm x 110mm, landscape) — 0.20
  Component 2: 25 sections (one per envelope) — 0.20
  Component 3: Company logo image present in envelopes — 0.15
  Component 4: Return address on envelopes — 0.15
  Component 5: All 25 vendor recipients from CSV present — 0.30
"""

import os
import csv

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_013'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------------------------------------------------------------
    # Component 1: DL envelope page dimensions (0.20 points)
    #   DL envelope = 220mm x 110mm. In landscape orientation,
    #   page_width ~ 220mm (7,920,000 EMU) and page_height ~ 110mm (3,960,000 EMU).
    #   Tolerance: +/- 5mm (180,000 EMU).
    # ---------------------------------------------------------------
    try:
        section = doc.sections[0]
        w = section.page_width
        h = section.page_height

        # Expected: width ~220mm = 7,920,000 EMU, height ~110mm = 3,960,000 EMU
        TOLERANCE = 180000  # ~5mm
        DL_WIDTH = 7920000
        DL_HEIGHT = 3960000

        width_ok = abs(w - DL_WIDTH) < TOLERANCE
        height_ok = abs(h - DL_HEIGHT) < TOLERANCE

        # Check landscape orientation
        from docx.enum.section import WD_ORIENT
        landscape_ok = section.orientation == WD_ORIENT.LANDSCAPE

        if width_ok and height_ok and landscape_ok:
            print(f"PASS: Component 1 — DL envelope dimensions: {w/36000:.1f}mm x {h/36000:.1f}mm, landscape (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Expected ~220mm x ~110mm landscape, got {w/36000:.1f}mm x {h/36000:.1f}mm, orient={section.orientation}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------------------------------------------------------------
    # Component 2: 25 sections (one per envelope) (0.20 points)
    #   The merged document should have exactly 25 sections, one per vendor.
    #   Partial credit: proportional to number of sections if >= 2.
    # ---------------------------------------------------------------
    try:
        num_sections = len(doc.sections)
        if num_sections >= 25:
            print(f"PASS: Component 2 — {num_sections} sections found (>= 25) (0.20 pts)")
            total_score += 0.20
        elif num_sections >= 2:
            partial = 0.20 * (num_sections / 25.0)
            print(f"PARTIAL: Component 2 — {num_sections}/25 sections ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {num_sections} section(s), expected 25")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---------------------------------------------------------------
    # Component 3: Company logo image present in envelopes (0.15 points)
    #   The golden doc has an image embedded (via doc.part.rels) and
    #   25 paragraphs referencing it (one per envelope).
    #   We check that at least one image exists AND at least 10 paragraphs
    #   contain image references (to confirm it's on multiple envelopes).
    # ---------------------------------------------------------------
    try:
        # Count image rels in document
        img_rels = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)

        # Count paragraphs that contain image references
        img_para_count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if 'graphicData' in run._element.xml:
                    img_para_count += 1
                    break  # one image per para is enough

        if img_rels >= 1 and img_para_count >= 10:
            print(f"PASS: Component 3 — {img_rels} image(s), {img_para_count} paragraphs with images (0.15 pts)")
            total_score += 0.15
        elif img_rels >= 1:
            # Image exists but not on enough envelopes
            partial = 0.075
            print(f"PARTIAL: Component 3 — Image found but only {img_para_count} paragraphs reference it ({partial:.3f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No images found in document")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ---------------------------------------------------------------
    # Component 4: Return address present on envelopes (0.15 points)
    #   The return address block contains "NEXUS Global Supply Co." repeated
    #   on each envelope. Check that it appears at least 20 times.
    # ---------------------------------------------------------------
    try:
        return_addr_count = 0
        for para in doc.paragraphs:
            if 'NEXUS Global Supply Co' in para.text:
                return_addr_count += 1

        if return_addr_count >= 20:
            print(f"PASS: Component 4 — Return address found {return_addr_count} times (0.15 pts)")
            total_score += 0.15
        elif return_addr_count >= 1:
            partial = 0.15 * min(return_addr_count / 20.0, 1.0)
            print(f"PARTIAL: Component 4 — Return address found {return_addr_count} times ({partial:.3f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Return address 'NEXUS Global Supply Co.' not found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ---------------------------------------------------------------
    # Component 5: All 25 vendor recipients from CSV present (0.30 points)
    #   Read vendors.csv and check that each VendorName appears in the document.
    #   Partial credit proportional to how many are found.
    # ---------------------------------------------------------------
    try:
        csv_path = os.path.join(WORKDIR, 'vendors.csv')
        if not os.path.exists(csv_path):
            print(f"FAIL: Component 5 — vendors.csv not found at {csv_path}")
        else:
            with open(csv_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                vendors = list(reader)

            # Collect all paragraph text for searching
            all_text = [para.text.strip() for para in doc.paragraphs]

            found_count = 0
            for vendor in vendors:
                vname = vendor.get('VendorName', '').strip()
                if vname and any(vname in t for t in all_text):
                    found_count += 1

            total_vendors = len(vendors)
            if found_count >= total_vendors and total_vendors >= 25:
                print(f"PASS: Component 5 — All {found_count}/{total_vendors} vendors found (0.30 pts)")
                total_score += 0.30
            elif found_count > 0:
                partial = 0.30 * (found_count / max(total_vendors, 25))
                print(f"PARTIAL: Component 5 — {found_count}/{total_vendors} vendors found ({partial:.3f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 — No vendor names found in document")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice state
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
