"""
Initial Setup: Insert page number field with offset starting from page 5
Task ID: writer_rd_063
Domain: libreoffice_writer

Creates a 20-page Writer document with front matter (pages 1-4) and main content
(pages 5-20). All pages show sequential page numbers 1-20 in the footer.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_063'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_field(paragraph):
    """Add a PAGE field code to a paragraph for automatic page numbering."""
    run_begin = paragraph.add_run()
    fld_begin = run_begin._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run_begin._element.append(fld_begin)

    run_instr = paragraph.add_run()
    instr = run_instr._element.makeelement(qn('w:instrText'), {})
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' PAGE '
    run_instr._element.append(instr)

    run_end = paragraph.add_run()
    fld_end = run_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run_end._element.append(fld_end)


def create_initial():
    doc = Document()

    # Set default page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Enable footer with page numbers for the default section
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fp.add_run("Page ")
    add_page_number_field(fp)

    # ==========================================
    # PAGE 1: Title Page
    # ==========================================
    title = doc.add_heading("The Art of Modern Software Architecture", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("")

    subtitle = doc.add_paragraph("A Comprehensive Guide to Building Scalable Systems")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.italic = True

    doc.add_paragraph("")
    author = doc.add_paragraph("By Dr. Elena Vasquez")
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in author.runs:
        run.font.size = Pt(14)

    edition = doc.add_paragraph("Third Edition, 2025")
    edition.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    publisher = doc.add_paragraph("Meridian Technical Press")
    publisher.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in publisher.runs:
        run.font.size = Pt(11)

    # ==========================================
    # PAGE 2: Dedication
    # ==========================================
    doc.add_page_break()
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    ded = doc.add_paragraph("To my mentors at MIT and Stanford,")
    ded.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in ded.runs:
        run.italic = True
        run.font.size = Pt(12)

    ded2 = doc.add_paragraph("who taught me that elegance in code")
    ded2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in ded2.runs:
        run.italic = True

    ded3 = doc.add_paragraph("is as important as functionality.")
    ded3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in ded3.runs:
        run.italic = True

    doc.add_paragraph("")
    ded4 = doc.add_paragraph("And to Alejandro, Sofia, and Mateo,")
    ded4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in ded4.runs:
        run.italic = True

    ded5 = doc.add_paragraph("for their endless patience and support.")
    ded5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in ded5.runs:
        run.italic = True

    # ==========================================
    # PAGE 3: Preface
    # ==========================================
    doc.add_page_break()
    doc.add_heading("Preface", level=1)

    doc.add_paragraph(
        "When I began writing the first edition of this book in 2019, the software industry "
        "was undergoing a fundamental transformation. Microservices architecture was replacing "
        "monolithic systems, cloud-native development was becoming the standard, and the tools "
        "available to architects were evolving at an unprecedented pace."
    )

    doc.add_paragraph(
        "This third edition reflects the maturation of many practices that were experimental "
        "just five years ago. Event-driven architecture, serverless computing, and AI-assisted "
        "development have moved from novelty to necessity. The chapters on distributed systems "
        "have been completely rewritten to incorporate lessons learned from large-scale deployments "
        "at companies like Netflix, Spotify, and Stripe."
    )

    doc.add_paragraph(
        "I am deeply grateful to the hundreds of practitioners who shared their experiences "
        "and case studies, and to the technical reviewers whose insights sharpened every chapter. "
        "Special thanks to Dr. James Park, Priya Sharma, and the engineering teams at CloudForge "
        "for their contributions to the new material on resilience patterns."
    )

    doc.add_paragraph(
        "My hope is that this book serves both as a practical guide for working architects "
        "and as an inspiration for those beginning their journey in software design."
    )

    preface_sig = doc.add_paragraph("Dr. Elena Vasquez")
    preface_sig.paragraph_format.space_before = Pt(24)
    sig2 = doc.add_paragraph("San Francisco, January 2025")

    # ==========================================
    # PAGE 4: Table of Contents
    # ==========================================
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)

    toc_entries = [
        ("Part I: Foundations", ""),
        ("  Chapter 1: Principles of Software Architecture", "5"),
        ("  Chapter 2: Architectural Patterns Overview", "7"),
        ("  Chapter 3: Quality Attributes and Trade-offs", "9"),
        ("Part II: Core Patterns", ""),
        ("  Chapter 4: Layered Architecture", "11"),
        ("  Chapter 5: Event-Driven Systems", "13"),
        ("  Chapter 6: Microservices at Scale", "15"),
        ("Part III: Advanced Topics", ""),
        ("  Chapter 7: Resilience and Fault Tolerance", "17"),
        ("  Chapter 8: Performance Optimization", "19"),
    ]

    for entry, page in toc_entries:
        p = doc.add_paragraph()
        if page:
            p.add_run(f"{entry} {'.' * (50 - len(entry))} {page}")
        else:
            r = p.add_run(entry)
            r.bold = True

    # ==========================================
    # PAGES 5-20: Main Content (16 pages)
    # ==========================================

    chapters = [
        {
            "title": "Chapter 1: Principles of Software Architecture",
            "content": [
                "Software architecture is the set of high-level decisions that shape the structure, behavior, and quality attributes of a system. Unlike detailed design, which focuses on algorithms and data structures, architecture addresses the fundamental organization of components and their interactions.",
                "The discipline of software architecture emerged in the 1990s as systems grew beyond the complexity that a single developer could manage. Pioneers like Mary Shaw, David Garlan, and Philippe Kruchten established the foundational vocabulary that practitioners use today.",
                "At its core, architecture is about making decisions that are costly to change later. These include the choice of programming paradigm, the communication protocol between services, the data storage strategy, and the deployment topology. Each decision constrains future options while enabling certain capabilities.",
                "Modern architecture practice emphasizes fitness functions: automated checks that verify a system continues to meet its architectural goals as it evolves. These functions complement traditional code reviews by providing continuous, objective feedback on architectural compliance.",
            ]
        },
        {
            "title": "Chapter 2: Architectural Patterns Overview",
            "content": [
                "Architectural patterns are proven structural templates that address recurring design challenges. They provide a common vocabulary for architects and developers, enabling teams to communicate complex ideas efficiently.",
                "The most widely recognized patterns include layered architecture, which separates concerns into horizontal tiers; client-server, which divides responsibilities between service providers and consumers; and pipe-and-filter, which processes data through a sequence of independent transformations.",
                "Event-driven architecture has gained enormous traction in recent years, particularly for systems that must respond to real-time changes. In this pattern, components communicate by producing and consuming events, enabling loose coupling and high scalability.",
                "The selection of an appropriate pattern depends on the system's quality attribute requirements. A healthcare records system prioritizing data consistency might favor a layered approach, while a social media platform emphasizing responsiveness might choose event-driven design.",
            ]
        },
        {
            "title": "Chapter 3: Quality Attributes and Trade-offs",
            "content": [
                "Quality attributes, also known as non-functional requirements, define how well a system performs its functions. They include performance, scalability, availability, security, maintainability, and testability, among others.",
                "The architecture of any significant system involves trade-offs between competing quality attributes. Improving security often reduces performance. Maximizing availability may increase costs. Enhancing flexibility can complicate the system's structure.",
                "The Architecture Tradeoff Analysis Method (ATAM) provides a systematic framework for evaluating these trade-offs. Developed at the Software Engineering Institute, ATAM helps stakeholders understand the consequences of architectural decisions before implementation begins.",
                "In practice, successful architects develop an intuition for trade-offs through experience. They learn to identify the two or three quality attributes most critical to a given system and optimize for those, while keeping others at acceptable levels.",
            ]
        },
        {
            "title": "Chapter 4: Layered Architecture",
            "content": [
                "Layered architecture organizes a system into horizontal tiers, each providing services to the tier above and consuming services from the tier below. The most common variant uses four layers: presentation, business logic, persistence, and database.",
                "The primary advantage of layering is separation of concerns. Changes to the user interface do not affect business rules. Database migrations do not require modifications to the presentation layer. This isolation simplifies maintenance and enables parallel development by different teams.",
                "However, layered architecture has significant drawbacks. The strict separation can lead to unnecessary data transformation as information passes through each layer. Performance suffers when every request must traverse all tiers, even when a direct path would be more efficient.",
                "Modern variants like the hexagonal architecture (ports and adapters) address these limitations by organizing dependencies around the domain model rather than technical layers. This approach keeps the core business logic independent of external concerns like databases and UI frameworks.",
            ]
        },
        {
            "title": "Chapter 5: Event-Driven Systems",
            "content": [
                "Event-driven architecture (EDA) models system behavior as a series of events: significant changes in state that other components may need to know about. An order placed, a payment received, a sensor reading exceeding a threshold are all events in this paradigm.",
                "The two primary topologies for event processing are mediator and broker. In the mediator topology, a central orchestrator coordinates event handling across multiple processors. In the broker topology, events flow through a distributed message broker with no central controller.",
                "Apache Kafka has become the de facto standard for event streaming at scale, handling millions of events per second with low latency and high durability. Its log-based architecture provides natural ordering guarantees and enables event replay for debugging and recovery.",
                "Designing effective event schemas requires careful attention to versioning, backward compatibility, and the distinction between events (something happened) and commands (do something). The CloudEvents specification provides a vendor-neutral format for event metadata.",
            ]
        },
        {
            "title": "Chapter 6: Microservices at Scale",
            "content": [
                "Microservices architecture decomposes a system into small, independently deployable services, each owning its data and implementing a bounded context from domain-driven design. This approach enables organizational scalability, allowing multiple teams to work on different services simultaneously.",
                "The transition from monolith to microservices is one of the most common and most challenging architectural migrations. Sam Newman's strangler fig pattern provides a gradual migration path, where new functionality is built as microservices while legacy code is progressively replaced.",
                "Service mesh technology, exemplified by Istio and Linkerd, has emerged as the standard infrastructure layer for microservice communication. It handles cross-cutting concerns like mutual TLS, circuit breaking, load balancing, and observability without requiring changes to application code.",
                "Data management in a microservices environment requires the saga pattern for distributed transactions, event sourcing for audit trails, and CQRS (Command Query Responsibility Segregation) for optimizing read and write paths independently.",
            ]
        },
        {
            "title": "Chapter 7: Resilience and Fault Tolerance",
            "content": [
                "Resilience is the ability of a system to handle and recover from failures gracefully. In distributed systems, failures are not exceptional events but routine occurrences. Network partitions, hardware failures, and software bugs are inevitable at scale.",
                "The circuit breaker pattern, popularized by Michael Nygard in 'Release It!', prevents cascading failures by monitoring the error rate of downstream calls. When failures exceed a threshold, the circuit opens and subsequent requests fail immediately rather than waiting for timeouts.",
                "Chaos engineering, pioneered by Netflix's Simian Army, takes a proactive approach to resilience by deliberately injecting failures into production systems. By regularly testing failure scenarios, teams build confidence in their system's ability to withstand unexpected disruptions.",
                "Bulkhead isolation limits the impact of failures by partitioning system resources. Just as bulkheads in a ship prevent flooding from spreading, resource isolation ensures that a failure in one component does not exhaust resources needed by others.",
            ]
        },
        {
            "title": "Chapter 8: Performance Optimization",
            "content": [
                "Performance optimization in distributed systems requires a systematic approach grounded in measurement. Premature optimization, as Donald Knuth famously warned, is the root of all evil. Every optimization decision should be driven by profiling data, not assumptions.",
                "Caching is the most impactful performance technique in most architectures. Redis and Memcached provide sub-millisecond response times for frequently accessed data. Content delivery networks cache static assets at the network edge, reducing latency for geographically distributed users.",
                "Database query optimization remains critical regardless of architectural style. Proper indexing, query plan analysis, connection pooling, and read replica strategies can improve database performance by orders of magnitude without changing application code.",
                "At the application level, asynchronous processing decouples request handling from long-running operations. Message queues like RabbitMQ and Amazon SQS enable background processing of tasks like email sending, report generation, and data aggregation.",
            ]
        },
    ]

    for i, chapter in enumerate(chapters):
        # Page break before each chapter (except the first which already has one from TOC)
        doc.add_page_break()
        doc.add_heading(chapter["title"], level=1)
        for para_text in chapter["content"]:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.space_after = Pt(6)

        # Add filler to ensure each chapter takes about 2 pages
        if i < len(chapters) - 1:
            doc.add_paragraph("")
            filler = doc.add_paragraph(
                "The implications of these principles extend beyond individual systems to the "
                "broader ecosystem of software development. Organizations that embrace architectural "
                "thinking at all levels of their engineering culture consistently deliver more "
                "reliable, maintainable, and adaptable systems. The following chapter builds on "
                "these concepts to explore more advanced techniques and real-world case studies."
            )
            filler.paragraph_format.space_after = Pt(6)

            doc.add_paragraph(
                "Key takeaways from this chapter include the importance of explicit trade-off "
                "analysis, the value of automation in maintaining architectural integrity, and "
                "the recognition that architecture is a continuous practice rather than a one-time "
                "activity. These principles will recur throughout the remainder of this book as "
                "we examine specific patterns and their applications in production environments."
            )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
