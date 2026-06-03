"""
Reward Script: Page number offset — front matter without numbers, main content starting at 1
Task ID: writer_rd_063
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30) — Document has at least 2 sections (section break separating front matter)
  Component 2 (0.25) — Front-matter section footer has NO page number field
  Component 3 (0.25) — Main-content section footer HAS page number field with 'Page' prefix
  Component 4 (0.20) — Main-content section has pgNumType restart at 1
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_063'

NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)
    print(f"INFO: Document has {num_sections} section(s)")

    # Component 1: Document has at least 2 sections (0.30 points)
    # The task requires a section break between front matter (pages 1-4) and main content (page 5+).
    # Initial doc has only 1 section, so this checks for the task-introduced change.
    try:
        if num_sections >= 2:
            print(f"PASS: Component 1 — Document has {num_sections} sections (>= 2) (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 1 — Expected >= 2 sections, found {num_sections}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Gate: need at least 2 sections to check the remaining components
    if num_sections < 2:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Component 2: Front-matter section (section 0) footer has NO page number field (0.25 points)
    # In the initial doc, section 0 has a PAGE field in the footer.
    # In the golden doc, section 0 footer is empty (no PAGE field) — front matter shows no numbers.
    try:
        sec0_footer = doc.sections[0].footer
        sec0_has_page_field = False
        for fp in sec0_footer.paragraphs:
            instr_texts = fp._element.findall('.//w:instrText', NS)
            for instr in instr_texts:
                if instr.text and 'PAGE' in instr.text.upper():
                    sec0_has_page_field = True
                    break

        if not sec0_has_page_field:
            print(f"PASS: Component 2 — Section 0 footer has no PAGE field (front matter hidden) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — Section 0 footer still has a PAGE field (front matter should have no page numbers)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Main-content section (section 1) footer HAS a PAGE field (0.25 points)
    # The main-content section must have a footer with "Page" prefix and a PAGE field code.
    try:
        sec1_footer = doc.sections[1].footer
        sec1_has_page_field = False
        sec1_has_page_prefix = False
        for fp in sec1_footer.paragraphs:
            # Check for PAGE field code
            instr_texts = fp._element.findall('.//w:instrText', NS)
            for instr in instr_texts:
                if instr.text and 'PAGE' in instr.text.upper():
                    sec1_has_page_field = True
            # Check for "Page" text prefix
            if 'Page' in fp.text or 'page' in fp.text.lower():
                sec1_has_page_prefix = True

        if sec1_has_page_field:
            detail = "with 'Page' prefix" if sec1_has_page_prefix else "without 'Page' prefix"
            print(f"PASS: Component 3 — Section 1 footer has PAGE field ({detail}) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 3 — Section 1 footer missing PAGE field (main content should show page numbers)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Main-content section has pgNumType with start=1 (0.20 points)
    # This is the page number restart — pages 5+ show numbers starting at 1 instead of 5.
    # The initial doc has no pgNumType restart anywhere.
    try:
        sec1_sectPr = doc.sections[1]._sectPr
        pgNumType = sec1_sectPr.find(qn('w:pgNumType'))
        if pgNumType is not None:
            start_val = pgNumType.get(qn('w:start'))
            if start_val == '1':
                print(f"PASS: Component 4 — Section 1 pgNumType start=1 (page numbering restarts at 1) (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 4 — pgNumType start={start_val}, expected '1'")
        else:
            print(f"FAIL: Component 4 — No pgNumType element in section 1 (page restart not set)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist app state before verification (in case doc is open in LibreOffice)
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
