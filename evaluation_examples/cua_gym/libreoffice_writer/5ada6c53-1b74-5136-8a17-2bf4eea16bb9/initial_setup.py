"""
Initial Setup: Engineering compensation structure document with introductory text
Task ID: writer_hr_039
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_039'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Engineering Compensation Structure', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introductory paragraphs about compensation philosophy
    doc.add_paragraph()

    p1 = doc.add_paragraph()
    p1.add_run('Compensation Philosophy').bold = True
    p1.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        'At Meridian Technologies, we believe that competitive and transparent compensation '
        'is essential to attracting and retaining top engineering talent. Our compensation '
        'framework is designed to reflect market rates, internal equity, and the increasing '
        'scope of responsibility at each career level.'
    )

    doc.add_paragraph(
        'Our engineering ladder consists of five distinct levels, each with a defined salary '
        'range that accounts for geographic adjustments, experience, and performance. The '
        'midpoint of each range represents the target compensation for a fully proficient '
        'engineer at that level, while the minimum and maximum allow for growth within a role.'
    )

    p2 = doc.add_paragraph()
    p2.add_run('Key Principles').bold = True
    p2.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        'Market Competitiveness: Salary ranges are benchmarked annually against industry '
        'surveys from Radford, Levels.fyi, and Mercer to ensure we remain within the 60th '
        'to 75th percentile for total compensation.',
        style='List Bullet'
    )

    doc.add_paragraph(
        'Internal Equity: Overlapping ranges between adjacent levels ensure that high '
        'performers at a lower level are not disadvantaged compared to new hires at the '
        'next level.',
        style='List Bullet'
    )

    doc.add_paragraph(
        'Transparency: All employees have access to the full salary structure. We '
        'encourage open conversations about career growth and compensation expectations '
        'with engineering managers.',
        style='List Bullet'
    )

    doc.add_paragraph(
        'Performance-Based Progression: Movement within a salary band is tied to bi-annual '
        'performance reviews, with top performers eligible for accelerated progression '
        'toward the band maximum.',
        style='List Bullet'
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'The table below should outline the salary structure for each engineering level. '
        'Please create it with the appropriate columns and data as specified by the HR team.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
