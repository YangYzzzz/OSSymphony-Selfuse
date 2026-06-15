"""
Reward Script: Salary structure table for engineering department
Task ID: writer_hr_039
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25) - Table exists with correct dimensions (6 rows x 5 cols)
  Component 2 (0.25) - Header row has correct text and bold formatting
  Component 3 (0.30) - Salary data rows contain correct values for all 5 levels
  Component 4 (0.20) - All cells have borders
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_039'

# Expected data for the salary table
EXPECTED_HEADERS = ['Level', 'Title', 'Salary Range (Min)', 'Salary Range (Mid)', 'Salary Range (Max)']

EXPECTED_ROWS = [
    # (Level, Title, Min, Mid, Max)
    ('E1', 'Junior Engineer',    65000,  75000,  85000),
    ('E2', 'Engineer',           80000,  95000, 110000),
    ('E3', 'Senior Engineer',   105000, 122500, 140000),
    ('E4', 'Staff Engineer',    130000, 152500, 175000),
    ('E5', 'Principal Engineer', 160000, 190000, 220000),
]


def parse_salary(text):
    """Extract numeric value from salary string like '$65,000' or '$65K'."""
    if not text:
        return None
    cleaned = text.replace('$', '').replace(',', '').replace(' ', '').strip()
    # Handle 'K' suffix
    if cleaned.upper().endswith('K'):
        try:
            return float(cleaned[:-1]) * 1000
        except ValueError:
            return None
    try:
        return float(cleaned)
    except ValueError:
        return None


def check_cell_has_border(cell):
    """Check if a cell has borders defined (top, bottom, left, right)."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return False
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        return False
    required = ['top', 'bottom', 'left', 'right']
    for side in required:
        border_el = borders.find(qn(f'w:{side}'))
        if border_el is None:
            return False
        val = border_el.get(qn('w:val'))
        if val is None or val == 'none' or val == 'nil':
            return False
    return True


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with correct dimensions (0.25 points)
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 1 -- No tables found in document")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows >= 6 and num_cols >= 5:
                print(f"PASS: Component 1 -- Table found with {num_rows} rows x {num_cols} cols (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 -- Table has {num_rows} rows x {num_cols} cols, expected >= 6x5")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Early exit if no table
    if len(doc.tables) == 0:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    table = doc.tables[0]

    # Component 2: Header row has correct text and bold formatting (0.25 points)
    try:
        header_row = table.rows[0]
        header_texts = [cell.text.strip() for cell in header_row.cells]

        # Check header text (0.10 pts)
        headers_correct = 0
        for i, expected in enumerate(EXPECTED_HEADERS):
            if i < len(header_texts) and expected.lower() in header_texts[i].lower():
                headers_correct += 1

        if headers_correct >= 4:
            print(f"PASS: Component 2a -- Header text matches ({headers_correct}/5 correct) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2a -- Header text mismatch ({headers_correct}/5). Found: {header_texts}")

        # Check header bold (0.15 pts)
        bold_count = 0
        for cell in header_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip() and run.font.bold:
                        bold_count += 1
                        break

        if bold_count >= 4:
            print(f"PASS: Component 2b -- Header row is bold ({bold_count}/5 cells bold) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2b -- Header row not fully bold ({bold_count}/5 cells bold)")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Salary data rows have correct values (0.30 points)
    try:
        rows_correct = 0
        pts_per_row = 0.06  # 0.06 * 5 = 0.30

        for row_idx, (exp_level, exp_title, exp_min, exp_mid, exp_max) in enumerate(EXPECTED_ROWS):
            data_row_idx = row_idx + 1  # skip header
            if data_row_idx >= len(table.rows):
                print(f"FAIL: Component 3 row {data_row_idx} -- Row missing")
                continue

            row = table.rows[data_row_idx]
            cells = [cell.text.strip() for cell in row.cells]

            # Check level
            level_ok = len(cells) > 0 and exp_level.lower() in cells[0].lower()
            # Check title
            title_ok = len(cells) > 1 and exp_title.lower() in cells[1].lower()
            # Check min salary
            min_val = parse_salary(cells[2]) if len(cells) > 2 else None
            min_ok = min_val is not None and abs(min_val - exp_min) < 1000
            # Check max salary
            max_val = parse_salary(cells[4]) if len(cells) > 4 else None
            max_ok = max_val is not None and abs(max_val - exp_max) < 1000
            # Check mid salary (midpoint)
            mid_val = parse_salary(cells[3]) if len(cells) > 3 else None
            mid_ok = mid_val is not None and abs(mid_val - exp_mid) < 1000

            if level_ok and title_ok and min_ok and max_ok and mid_ok:
                print(f"PASS: Component 3 row {data_row_idx} -- {exp_level} {exp_title} data correct ({pts_per_row} pts)")
                total_score += pts_per_row
                rows_correct += 1
            else:
                details = f"level={level_ok}, title={title_ok}, min={min_ok}({min_val}), mid={mid_ok}({mid_val}), max={max_ok}({max_val})"
                print(f"FAIL: Component 3 row {data_row_idx} -- {exp_level} checks: {details}")

        print(f"Component 3 summary: {rows_correct}/5 rows correct")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: All cells have borders (0.20 points)
    try:
        total_cells = 0
        cells_with_borders = 0
        for row in table.rows:
            for cell in row.cells:
                total_cells += 1
                if check_cell_has_border(cell):
                    cells_with_borders += 1

        border_ratio = cells_with_borders / total_cells if total_cells > 0 else 0

        if border_ratio >= 0.9:
            print(f"PASS: Component 4 -- {cells_with_borders}/{total_cells} cells have borders (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 -- Only {cells_with_borders}/{total_cells} cells have borders (ratio={border_ratio:.2f})")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
