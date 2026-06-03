"""
Initial Setup: Accept/reject tracked changes in a business proposal
Task ID: writer_rm_034
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from lxml import etree
from copy import deepcopy
import datetime

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_034'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# OOXML namespaces
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
nsmap = {'w': W_NS, 'r': R_NS}


def make_run_element(text, bold=False, italic=False, font_name=None, font_size_pt=None, color_rgb=None):
    """Create a w:r element with optional formatting."""
    r = etree.SubElement(etree.Element('dummy'), f'{{{W_NS}}}r')
    rPr = etree.SubElement(r, f'{{{W_NS}}}rPr')
    if bold:
        etree.SubElement(rPr, f'{{{W_NS}}}b')
    if italic:
        etree.SubElement(rPr, f'{{{W_NS}}}i')
    if font_name:
        rFonts = etree.SubElement(rPr, f'{{{W_NS}}}rFonts')
        rFonts.set(f'{{{W_NS}}}ascii', font_name)
        rFonts.set(f'{{{W_NS}}}hAnsi', font_name)
    if font_size_pt:
        sz = etree.SubElement(rPr, f'{{{W_NS}}}sz')
        sz.set(f'{{{W_NS}}}val', str(font_size_pt * 2))  # half-points
    if color_rgb:
        color_el = etree.SubElement(rPr, f'{{{W_NS}}}color')
        color_el.set(f'{{{W_NS}}}val', color_rgb)
    t = etree.SubElement(r, f'{{{W_NS}}}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    return r


def make_del_run_element(text, bold=False, italic=False, font_name=None, font_size_pt=None):
    """Create a w:r element with w:delText instead of w:t (for deleted text)."""
    r = etree.SubElement(etree.Element('dummy'), f'{{{W_NS}}}r')
    rPr = etree.SubElement(r, f'{{{W_NS}}}rPr')
    if bold:
        etree.SubElement(rPr, f'{{{W_NS}}}b')
    if italic:
        etree.SubElement(rPr, f'{{{W_NS}}}i')
    if font_name:
        rFonts = etree.SubElement(rPr, f'{{{W_NS}}}rFonts')
        rFonts.set(f'{{{W_NS}}}ascii', font_name)
        rFonts.set(f'{{{W_NS}}}hAnsi', font_name)
    if font_size_pt:
        sz = etree.SubElement(rPr, f'{{{W_NS}}}sz')
        sz.set(f'{{{W_NS}}}val', str(font_size_pt * 2))
    dt = etree.SubElement(r, f'{{{W_NS}}}delText')
    dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    dt.text = text
    return r


def add_tracked_insertion(parent_para, text, rev_id, author="Sarah Chen",
                          date="2025-11-15T09:30:00Z", bold=False, italic=False,
                          font_name=None, font_size_pt=None, position=None):
    """Add a tracked insertion (w:ins) to a paragraph element."""
    ins = etree.SubElement(etree.Element('dummy'), f'{{{W_NS}}}ins')
    ins.set(f'{{{W_NS}}}id', str(rev_id))
    ins.set(f'{{{W_NS}}}author', author)
    ins.set(f'{{{W_NS}}}date', date)
    run = make_run_element(text, bold=bold, italic=italic, font_name=font_name, font_size_pt=font_size_pt)
    ins.append(run)
    if position is not None:
        parent_para.insert(position, ins)
    else:
        parent_para.append(ins)
    return ins


def add_tracked_deletion(parent_para, text, rev_id, author="Sarah Chen",
                         date="2025-11-15T09:30:00Z", bold=False, italic=False,
                         font_name=None, font_size_pt=None, position=None):
    """Add a tracked deletion (w:del) to a paragraph element."""
    del_el = etree.SubElement(etree.Element('dummy'), f'{{{W_NS}}}del')
    del_el.set(f'{{{W_NS}}}id', str(rev_id))
    del_el.set(f'{{{W_NS}}}author', author)
    del_el.set(f'{{{W_NS}}}date', date)
    run = make_del_run_element(text, bold=bold, italic=italic, font_name=font_name, font_size_pt=font_size_pt)
    del_el.append(run)
    if position is not None:
        parent_para.insert(position, del_el)
    else:
        parent_para.append(del_el)
    return del_el


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # First create the base document with python-docx for proper styling
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ======= SECTION 1: Overview =======
    h1 = doc.add_heading('Section 1: Overview', level=1)

    # Paragraph with tracked change #1: title rewording
    # Original: "Project Alpha: A Strategic Initiative"
    # Changed to: "Project Alpha: A Comprehensive Strategic Initiative"
    p1 = doc.add_paragraph()
    # We'll manipulate XML after save/load

    # Paragraph with tracked change #2: date update
    p2 = doc.add_paragraph()

    # Regular paragraph
    p3 = doc.add_paragraph(
        "The project aims to modernize our customer relationship management system "
        "and integrate it with our existing enterprise resource planning platform. "
        "This will streamline operations across all departments."
    )

    # Paragraph with tracked change #3: added sentence
    p4 = doc.add_paragraph()

    p5 = doc.add_paragraph(
        "Key stakeholders include the VP of Engineering, Director of Operations, "
        "and the Chief Financial Officer. Regular progress reviews will be conducted "
        "on a bi-weekly basis."
    )

    # ======= SECTION 2: Budget =======
    h2 = doc.add_heading('Section 2: Budget', level=1)

    p6 = doc.add_paragraph(
        "The following budget allocations have been approved by the finance committee "
        "for the fiscal year 2025-2026."
    )

    # Budget table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Category', 'Allocated Budget', 'Notes']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    budget_data = [
        ['Personnel', '$245,000', 'Full-time and contract staff'],
        ['Infrastructure', '$128,500', 'Cloud hosting and hardware'],
        ['Software Licenses', '$67,200', 'Annual subscriptions'],
        ['Training & Development', '$34,800', 'Staff certification programs'],
        ['Contingency Fund', '$25,000', 'Emergency reserves'],
    ]
    for r, row_data in enumerate(budget_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # Paragraphs with tracked changes #4-7 (budget reductions)
    p7 = doc.add_paragraph()
    p8 = doc.add_paragraph()
    p9 = doc.add_paragraph()
    p10 = doc.add_paragraph(
        "All budget modifications require written approval from the CFO and "
        "must be submitted at least 30 days before the start of each quarter."
    )

    # ======= SECTION 3: Timeline =======
    h3 = doc.add_heading('Section 3: Timeline', level=1)

    p11 = doc.add_paragraph(
        "The project timeline spans 18 months, divided into four major phases."
    )

    # Paragraph with tracked changes #8-9 (untouched)
    p12 = doc.add_paragraph()
    p13 = doc.add_paragraph()

    p14 = doc.add_paragraph(
        "Final deliverables are expected by December 2026. A comprehensive "
        "post-implementation review will follow in Q1 2027."
    )

    doc.save(OUTPUT)

    # Now reopen with lxml to inject tracked changes at XML level
    doc2 = Document(OUTPUT)
    body = doc2.element.body

    # Get all paragraphs
    paras = body.findall(f'{{{W_NS}}}p')
    # Also include table paragraphs - but let's index body-level elements

    # We need to find the right paragraphs. Let's list them:
    # Index 0: heading "Section 1: Overview"
    # Index 1: p1 (empty - for tracked change #1)
    # Index 2: p2 (empty - for tracked change #2)
    # Index 3: p3 (regular text about CRM)
    # Index 4: p4 (empty - for tracked change #3)
    # Index 5: p5 (regular about stakeholders)
    # Index 6: heading "Section 2: Budget"
    # Index 7: p6 (budget intro)
    # -- table is not a <w:p>, it's a <w:tbl> --
    # Index 8: p7 (empty - tracked change #4)
    # Index 9: p8 (empty - tracked change #5)
    # Index 10: p9 (empty - tracked change #6)
    # Index 11: p10 (budget approval text) - tracked change #7 will be inside this
    # Index 12: heading "Section 3: Timeline"
    # Index 13: p11 (timeline text)
    # Index 14: p12 (empty - tracked change #8)
    # Index 15: p13 (empty - tracked change #9)
    # Index 16: p14 (final deliverables)

    rev_id = 1
    author_s = "Sarah Chen"
    author_m = "Marcus Rivera"
    date1 = "2025-11-12T14:22:00Z"
    date2 = "2025-11-13T10:15:00Z"

    # ====== SECTION 1 TRACKED CHANGES (3 changes, to be accepted) ======

    # Change #1 (para index 1): Title rewording
    # Original: "Project Alpha: A Strategic Initiative"
    # New: "Project Alpha: A Comprehensive Strategic Initiative"
    p1_el = paras[1]
    # Add the normal part first
    r_normal = make_run_element("Project Alpha: A ", bold=True, font_size_pt=13)
    p1_el.append(r_normal)
    # Delete old word "Strategic"
    add_tracked_deletion(p1_el, "Strategic", rev_id, author=author_s, date=date1, bold=True, font_size_pt=13)
    rev_id += 1
    # Insert new word "Comprehensive Strategic"
    add_tracked_insertion(p1_el, "Comprehensive Strategic", rev_id, author=author_s, date=date1, bold=True, font_size_pt=13)
    rev_id += 1
    # Normal trailing text
    r_trail = make_run_element(" Initiative", bold=True, font_size_pt=13)
    p1_el.append(r_trail)

    # Change #2 (para index 2): Date update
    # Original: "Proposed start date: January 15, 2026"
    # Changed to: "Proposed start date: March 1, 2026"
    p2_el = paras[2]
    r_date_pre = make_run_element("Proposed start date: ")
    p2_el.append(r_date_pre)
    add_tracked_deletion(p2_el, "January 15, 2026", rev_id, author=author_s, date=date1)
    rev_id += 1
    add_tracked_insertion(p2_el, "March 1, 2026", rev_id, author=author_s, date=date1)
    rev_id += 1

    # Change #3 (para index 4): Added sentence
    # New sentence inserted: "An executive summary will be distributed to all department heads prior to the kickoff meeting."
    p4_el = paras[4]
    r_existing = make_run_element(
        "Resource allocation will be finalized after the initial planning phase is complete. "
    )
    p4_el.append(r_existing)
    add_tracked_insertion(
        p4_el,
        "An executive summary will be distributed to all department heads prior to the kickoff meeting.",
        rev_id, author=author_s, date=date2
    )
    rev_id += 1

    # ====== SECTION 2 TRACKED CHANGES (4 changes, to be rejected) ======
    # These are cost reduction proposals that were NOT approved

    # Change #4 (para index 8): Reduce personnel budget
    # Original: "The personnel budget should remain at $245,000 as approved."
    # Changed to: "The personnel budget should be reduced to $198,000 to cut costs."
    p7_el = paras[8]
    r_pre4 = make_run_element("The personnel budget should ")
    p7_el.append(r_pre4)
    add_tracked_deletion(p7_el, "remain at $245,000 as approved.", rev_id, author=author_m, date=date2)
    rev_id += 1
    add_tracked_insertion(p7_el, "be reduced to $198,000 to cut costs.", rev_id, author=author_m, date=date2)
    rev_id += 1

    # Change #5 (para index 9): Reduce infrastructure budget
    # Original: "Infrastructure spending of $128,500 covers essential cloud migration."
    # Changed to: "Infrastructure spending of $95,000 covers essential cloud migration."
    p8_el = paras[9]
    r_pre5 = make_run_element("Infrastructure spending of ")
    p8_el.append(r_pre5)
    add_tracked_deletion(p8_el, "$128,500", rev_id, author=author_m, date=date2)
    rev_id += 1
    add_tracked_insertion(p8_el, "$95,000", rev_id, author=author_m, date=date2)
    rev_id += 1
    r_post5 = make_run_element(" covers essential cloud migration.")
    p8_el.append(r_post5)

    # Change #6 (para index 10): Remove training budget line
    # Original: "Training programs are critical for successful adoption of the new system."
    # Changed to: "Training programs can be deferred to the next fiscal year."
    p9_el = paras[10]
    r_pre6 = make_run_element("Training programs ")
    p9_el.append(r_pre6)
    add_tracked_deletion(p9_el, "are critical for successful adoption of the new system.", rev_id, author=author_m, date=date2)
    rev_id += 1
    add_tracked_insertion(p9_el, "can be deferred to the next fiscal year.", rev_id, author=author_m, date=date2)
    rev_id += 1

    # Change #7 (para index 11): Modify approval text
    # This paragraph already has text, we need to clear it and rebuild with tracked change
    # Original: "All budget modifications require written approval from the CFO and must be submitted at least 30 days before the start of each quarter."
    # Changed to: "All budget modifications require email approval from the department manager and must be submitted at least 15 days before the start of each quarter."
    p10_el = paras[11]
    # Clear existing content (runs) but keep pPr if any
    pPr = p10_el.find(f'{{{W_NS}}}pPr')
    for child in list(p10_el):
        if child.tag != f'{{{W_NS}}}pPr':
            p10_el.remove(child)

    r_pre7 = make_run_element("All budget modifications require ")
    p10_el.append(r_pre7)
    add_tracked_deletion(p10_el, "written approval from the CFO", rev_id, author=author_m, date=date2)
    rev_id += 1
    add_tracked_insertion(p10_el, "email approval from the department manager", rev_id, author=author_m, date=date2)
    rev_id += 1
    r_mid7 = make_run_element(" and must be submitted at least ")
    p10_el.append(r_mid7)
    add_tracked_deletion(p10_el, "30 days", rev_id, author=author_m, date=date2)
    rev_id += 1
    add_tracked_insertion(p10_el, "15 days", rev_id, author=author_m, date=date2)
    rev_id += 1
    r_post7 = make_run_element(" before the start of each quarter.")
    p10_el.append(r_post7)

    # ====== SECTION 3 TRACKED CHANGES (2 changes, to remain untouched) ======

    # Change #8 (para index 14): Phase name change
    # Original: "Phase 2 is scheduled for Q2 2026 and focuses on system integration."
    # Changed to: "Phase 2 is scheduled for Q2 2026 and focuses on data migration and system integration."
    p12_el = paras[14]
    r_pre8 = make_run_element("Phase 2 is scheduled for Q2 2026 and focuses on ")
    p12_el.append(r_pre8)
    add_tracked_insertion(p12_el, "data migration and ", rev_id, author=author_s, date=date2)
    rev_id += 1
    r_post8 = make_run_element("system integration.")
    p12_el.append(r_post8)

    # Change #9 (para index 15): Add risk note
    # Original: "Phase 3 covers user acceptance testing during Q3 2026."
    # Changed to: "Phase 3 covers user acceptance testing and risk assessment during Q3 2026."
    p13_el = paras[15]
    r_pre9 = make_run_element("Phase 3 covers user acceptance testing ")
    p13_el.append(r_pre9)
    add_tracked_insertion(p13_el, "and risk assessment ", rev_id, author=author_s, date=date2)
    rev_id += 1
    r_post9 = make_run_element("during Q3 2026.")
    p13_el.append(r_post9)

    doc2.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
