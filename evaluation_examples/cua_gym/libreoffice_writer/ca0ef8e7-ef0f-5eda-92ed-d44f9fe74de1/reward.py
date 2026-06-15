"""
Reward Script: Create invoice table with merged header, line items, and SUM formula
Task ID: writer_tm_050
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table exists with 9 rows and 4 columns
  Component 2 (0.25): Row 0 cells merged across all 4 columns (gridSpan=4)
  Component 3 (0.20): Row 1 contains headers Item/Description/Qty/Price
  Component 4 (0.30): Row 8 has 'Total' in col A and SUM formula in col D
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_050'


def persist_app_state(domain: str):
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Cannot import python-docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: at least one table must exist
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document — task requires creating a table")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # Component 1: Table has 9 rows and 4 columns (0.25 points)
    try:
        if num_rows == 9 and num_cols == 4:
            print(f"PASS: Component 1 — Table has {num_rows} rows and {num_cols} columns (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Expected 9 rows x 4 cols, found {num_rows} rows x {num_cols} cols")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Row 0 is merged across all 4 columns (gridSpan=4) (0.25 points)
    try:
        row0_cell0 = table.rows[0].cells[0]
        tc_pr = row0_cell0._tc.find(qn('w:tcPr'))
        grid_span = None
        if tc_pr is not None:
            gs_elem = tc_pr.find(qn('w:gridSpan'))
            if gs_elem is not None:
                grid_span = int(gs_elem.get(qn('w:val')))

        if grid_span is not None and grid_span >= 4:
            print(f"PASS: Component 2 — Row 0 merged with gridSpan={grid_span} (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — Row 0 not merged across 4 columns (gridSpan={grid_span})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Row 1 has headers Item, Description, Qty, Price (0.20 points)
    try:
        if num_rows >= 2:
            row1_texts = [table.cell(1, j).text.strip().lower() for j in range(min(num_cols, 4))]
            expected_headers = ['item', 'description', 'qty', 'price']
            matches = sum(1 for a, b in zip(row1_texts, expected_headers) if a == b)
            if matches == 4:
                print(f"PASS: Component 3 — Headers match: {row1_texts} (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 — Expected {expected_headers}, found {row1_texts} ({matches}/4 match)")
        else:
            print(f"FAIL: Component 3 — Table has fewer than 2 rows")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Row 8 has 'Total' in col A and SUM formula in col D (0.30 points)
    try:
        if num_rows >= 9:
            # Check 'Total' in A9 (row index 8, col 0)
            a9_text = table.cell(8, 0).text.strip().lower()
            has_total = 'total' in a9_text

            # Check for SUM formula in D9 (row index 8, col 3)
            d9_cell = table.cell(8, 3)
            d9_xml = d9_cell._tc.xml
            has_sum = False

            # Check fldSimple with SUM
            import re
            fld_simple_matches = re.findall(r'fldSimple[^>]*instr="([^"]*)"', d9_xml)
            for instr in fld_simple_matches:
                if 'SUM' in instr.upper():
                    has_sum = True
                    break

            # Also check instrText (complex field codes)
            if not has_sum:
                instr_text_matches = re.findall(r'instrText[^>]*>([^<]+)<', d9_xml)
                for instr in instr_text_matches:
                    if 'SUM' in instr.upper():
                        has_sum = True
                        break

            if has_total and has_sum:
                print(f"PASS: Component 4 — 'Total' in A9 and SUM formula in D9 (0.30 pts)")
                total_score += 0.30
            elif has_total:
                print(f"PARTIAL: Component 4 — 'Total' found in A9 but no SUM formula in D9 (0.10 pts)")
                total_score += 0.10
            elif has_sum:
                print(f"PARTIAL: Component 4 — SUM formula found in D9 but 'Total' missing in A9 (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 4 — No 'Total' in A9 (found: '{a9_text}') and no SUM formula in D9")
        else:
            print(f"FAIL: Component 4 — Table has fewer than 9 rows ({num_rows})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
