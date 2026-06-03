"""
Initial Setup: Create a document with company header text, no table.
Task ID: writer_tm_050
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_050'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Company header as a heading
    heading = doc.add_heading('Acme Corp', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in heading.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Invoice')
    run.font.size = Pt(16)
    run.bold = True

    # Some introductory text
    doc.add_paragraph('')  # blank line

    intro = doc.add_paragraph()
    run = intro.add_run('Date: March 15, 2025')
    run.font.size = Pt(11)

    addr = doc.add_paragraph()
    run = addr.add_run('Bill To: Meridian Consulting Group')
    run.font.size = Pt(11)

    addr2 = doc.add_paragraph()
    run = addr2.add_run('1240 Oak Valley Drive, Suite 300, San Francisco, CA 94102')
    run.font.size = Pt(11)

    doc.add_paragraph('')  # blank line

    notes = doc.add_paragraph()
    run = notes.add_run(
        'Please add an invoice table below with line items and a total row. '
        'The table should include columns for Item, Description, Qty, and Price.'
    )
    run.font.size = Pt(11)
    run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
