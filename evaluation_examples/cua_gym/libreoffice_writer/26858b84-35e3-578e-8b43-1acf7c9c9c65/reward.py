"""
Reward Script: Add multiplication formulas in table column E (Unit Price * Qty)
Task ID: writer_tm_034
Domain: libreoffice_writer
Scoring: 5 components (0.2 each) — one per row E2-E6 checking correct product value
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_034'

# Expected products: C_col * D_col for each row
EXPECTED = {
    1: 100.00,   # 25.00 * 4
    2: 155.00,   # 15.50 * 10
    3: 84.00,    # 42.00 * 2
    4: 175.00,   # 8.75 * 20
    5: 165.00,   # 33.00 * 5
}


def persist_app_state(domain: str):
    """Best-effort save in case LibreOffice has unsaved changes."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def parse_number(text):
    """Parse a numeric value from cell text, tolerating formulas or plain numbers."""
    text = text.strip()
    if not text:
        return None
    # Remove currency symbols, commas
    cleaned = re.sub(r'[,$\s]', '', text)
    try:
        return float(cleaned)
    except ValueError:
        return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Precondition: table must have at least 6 rows and 5 columns
    if len(table.rows) < 6 or len(table.columns) < 5:
        print(f"FAIL: Table too small — {len(table.rows)} rows x {len(table.columns)} cols (need 6x5)")
        print("REWARD: 0.0")
        return 0.0

    # Components 1-5: Check E2-E6 values (rows 1-5, col 4)
    for row_idx in range(1, 6):
        component_num = row_idx
        expected_val = EXPECTED[row_idx]

        try:
            cell_text = table.cell(row_idx, 4).text.strip()
            actual_val = parse_number(cell_text)

            if actual_val is not None and abs(actual_val - expected_val) < 0.01:
                print(f"PASS: Component {component_num} — Row {row_idx+1} col E = {actual_val} (expected {expected_val}) (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component {component_num} — Row {row_idx+1} col E: got '{cell_text}' (parsed: {actual_val}), expected {expected_val}")
        except Exception as e:
            print(f"ERROR: Component {component_num} — Row {row_idx+1} col E: {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
