"""
Initial Setup: Create Order Form document with a table for formula task
Task ID: writer_tm_034
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_034'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading('Order Form', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle info
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run('Client: Riverside Electronics Ltd.')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    p2 = doc.add_paragraph()
    run2 = p2.add_run('Order Date: 2025-09-12')
    run2.font.size = Pt(11)
    run2.font.name = 'Calibri'

    p3 = doc.add_paragraph()
    run3 = p3.add_run('PO Number: PO-2025-4871')
    run3.font.size = Pt(11)
    run3.font.name = 'Calibri'

    doc.add_paragraph()  # spacing

    # Create 7-row x 5-col table (1 header + 5 data + 1 totals row)
    table = doc.add_table(rows=7, cols=5)
    table.style = 'Table Grid'

    # Headers
    headers = ['Item', 'Description', 'Unit Price', 'Qty', 'Total']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    # Data rows (rows 2-6 in task notation = table rows 1-5 in 0-indexed)
    data = [
        ['USB-C Hub',       'Premium 7-in-1 USB-C docking station',   '25.00',  '4',  ''],
        ['HDMI Cable 2m',   'High-speed HDMI 2.1 braided cable',      '15.50',  '10', ''],
        ['Wireless Mouse',  'Ergonomic Bluetooth mouse, rechargeable', '42.00',  '2',  ''],
        ['Cable Ties',      'Reusable velcro cable ties, pack of 50',  '8.75',   '20', ''],
        ['Monitor Stand',   'Adjustable aluminum monitor riser',       '33.00',  '5',  ''],
    ]

    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table.cell(i + 1, j)
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'

    # Row 7 (index 6) - Totals row placeholder
    total_label_cell = table.cell(6, 0)
    total_label_cell.text = ''
    # Merge first 4 cells for label
    merged = table.cell(6, 0).merge(table.cell(6, 3))
    run = merged.paragraphs[0].add_run('Grand Total')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    merged.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Leave E7 (total column in totals row) empty
    table.cell(6, 4).text = ''

    # Footer note
    doc.add_paragraph()
    p_note = doc.add_paragraph()
    run_note = p_note.add_run('Terms: Net 30 days. Prices in USD. Shipping charges apply separately.')
    run_note.font.size = Pt(9)
    run_note.font.name = 'Calibri'
    run_note.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
