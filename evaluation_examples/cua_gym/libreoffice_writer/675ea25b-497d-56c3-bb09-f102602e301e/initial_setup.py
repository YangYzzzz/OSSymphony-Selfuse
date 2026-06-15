"""
Initial Setup: Physics Textbook Master Document with per-chapter numbering
Task ID: writer_rm_095
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_095'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    return h


def add_caption(doc, caption_text):
    """Add a caption paragraph (e.g., 'Figure 1: Description')."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(caption_text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


def add_equation_block(doc, equation_text, label):
    """Add an equation-like block with label."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(equation_text)
    run.font.name = "Courier New"
    run.font.size = Pt(11)
    # Add equation label on the right
    run2 = p.add_run(f"    ({label})")
    run2.font.size = Pt(10)
    run2.italic = True
    return p


def add_data_table(doc, headers, rows, caption):
    """Add a data table with caption."""
    add_caption(doc, caption)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            table.cell(i + 1, j).text = str(val)
    doc.add_paragraph("")  # spacing


# ============================================================
# CHAPTER DATA: Each chapter has figures, tables, equations
# with numbering that RESTARTS at 1 per chapter
# ============================================================

chapters = [
    {
        "title": "Chapter 1: Classical Mechanics",
        "intro": "Classical mechanics describes the motion of macroscopic objects under the influence of forces. This chapter covers Newton's laws, kinematics, and energy conservation principles that form the foundation of physics.",
        "sections": [
            ("1.1 Newton's Laws of Motion", "Newton's three laws of motion provide the framework for understanding how forces affect the motion of objects. The first law states that an object at rest stays at rest unless acted upon by an external force. The second law quantifies the relationship between force, mass, and acceleration."),
            ("1.2 Work and Energy", "The work-energy theorem establishes the connection between net work done on an object and its change in kinetic energy. Conservation of energy is one of the most fundamental principles in all of physics."),
            ("1.3 Momentum and Collisions", "Linear momentum is conserved in isolated systems. This principle is particularly useful for analyzing collisions, whether elastic or inelastic."),
        ],
        "figures": [
            ("Figure 1: Free body diagram of a block on an inclined plane", "The figure shows force vectors including gravity (mg), normal force (N), and friction (f) acting on a 5 kg block resting on a 30-degree inclined surface."),
            ("Figure 2: Energy conservation in a pendulum system", "A simple pendulum of length L = 1.2 m demonstrating the conversion between potential and kinetic energy at various swing positions."),
        ],
        "tables": [
            {
                "caption": "Table 1: Kinematic equations for constant acceleration",
                "headers": ["Equation", "Variables", "Missing Variable"],
                "rows": [
                    ["v = v₀ + at", "v, v₀, a, t", "Δx"],
                    ["Δx = v₀t + ½at²", "Δx, v₀, a, t", "v"],
                    ["v² = v₀² + 2aΔx", "v, v₀, a, Δx", "t"],
                    ["Δx = ½(v₀ + v)t", "Δx, v₀, v, t", "a"],
                ],
            },
        ],
        "equations": [
            ("F = ma", "Equation 1"),
            ("W = F · d · cos(θ)", "Equation 2"),
            ("p = mv", "Equation 3"),
        ],
    },
    {
        "title": "Chapter 2: Thermodynamics",
        "intro": "Thermodynamics governs the behavior of heat, work, and energy in physical systems. The laws of thermodynamics define fundamental constraints on energy transformations and establish the concept of entropy.",
        "sections": [
            ("2.1 Heat and Temperature", "Temperature is a measure of the average kinetic energy of particles in a substance. Heat is the transfer of thermal energy between objects at different temperatures."),
            ("2.2 Laws of Thermodynamics", "The four laws of thermodynamics establish the rules governing energy transfer. The zeroth law defines thermal equilibrium, while the first law is conservation of energy applied to thermal systems."),
        ],
        "figures": [
            ("Figure 1: PV diagram for an ideal Carnot cycle", "The diagram shows four stages: isothermal expansion, adiabatic expansion, isothermal compression, and adiabatic compression for a gas with T_H = 600 K and T_C = 300 K."),
            ("Figure 2: Heat engine efficiency comparison", "Bar chart comparing theoretical maximum (Carnot) efficiency with actual efficiencies of gasoline engines (25%), diesel engines (40%), and steam turbines (45%)."),
            ("Figure 3: Entropy changes during phase transitions of water", "Graph showing entropy as a function of temperature for H₂O from -20°C to 150°C at 1 atm, with plateaus at phase transitions."),
        ],
        "tables": [
            {
                "caption": "Table 1: Specific heat capacities of common materials",
                "headers": ["Material", "c (J/kg·K)", "State"],
                "rows": [
                    ["Water", "4186", "Liquid"],
                    ["Ice", "2090", "Solid"],
                    ["Steam", "2010", "Gas"],
                    ["Aluminum", "900", "Solid"],
                    ["Copper", "385", "Solid"],
                    ["Iron", "449", "Solid"],
                ],
            },
        ],
        "equations": [
            ("Q = mcΔT", "Equation 1"),
            ("ΔU = Q - W", "Equation 2"),
            ("η = 1 - T_C/T_H", "Equation 3"),
            ("ΔS = Q/T", "Equation 4"),
        ],
    },
    {
        "title": "Chapter 3: Waves and Oscillations",
        "intro": "Wave phenomena pervade every branch of physics, from sound to light to quantum mechanics. This chapter explores the mathematical description of waves, superposition, and resonance.",
        "sections": [
            ("3.1 Simple Harmonic Motion", "Simple harmonic motion occurs when the restoring force is proportional to displacement. The mass-spring system and the simple pendulum are classic examples."),
            ("3.2 Wave Properties", "Waves transport energy without transporting matter. Key properties include wavelength, frequency, amplitude, and wave speed."),
        ],
        "figures": [
            ("Figure 1: Standing wave patterns on a vibrating string", "Diagram showing the first four harmonics of a string fixed at both ends, with nodes and antinodes labeled for L = 2.0 m."),
            ("Figure 2: Doppler effect for a moving sound source", "Illustration of wavefront compression ahead of a source moving at v_s = 30 m/s in air (v_sound = 343 m/s)."),
        ],
        "tables": [
            {
                "caption": "Table 1: Speed of sound in various media at 20°C",
                "headers": ["Medium", "Speed (m/s)", "Type"],
                "rows": [
                    ["Air", "343", "Gas"],
                    ["Helium", "1007", "Gas"],
                    ["Water", "1482", "Liquid"],
                    ["Steel", "5960", "Solid"],
                    ["Glass", "5640", "Solid"],
                ],
            },
        ],
        "equations": [
            ("x(t) = A cos(ωt + φ)", "Equation 1"),
            ("v = fλ", "Equation 2"),
            ("f' = f(v ± v_o)/(v ∓ v_s)", "Equation 3"),
        ],
    },
    {
        "title": "Chapter 4: Electrostatics",
        "intro": "Electrostatics is the study of electric charges at rest and the forces between them. Coulomb's law and Gauss's law provide the mathematical foundation for understanding electric fields and potentials.",
        "sections": [
            ("4.1 Electric Charge and Coulomb's Law", "Electric charge is a fundamental property of matter. Like charges repel and unlike charges attract, with the force inversely proportional to the square of the distance."),
            ("4.2 Electric Fields and Potential", "The electric field is a vector field that describes the force per unit charge at each point in space. Electric potential is the scalar field representing potential energy per unit charge."),
        ],
        "figures": [
            ("Figure 1: Electric field lines between two point charges", "Field line diagram for a +3 μC and -1 μC charge separated by 0.5 m, showing field concentration near the negative charge."),
            ("Figure 2: Equipotential surfaces around a positive point charge", "Concentric spherical surfaces at V = 100V, 50V, 25V around a +2 μC charge in free space."),
        ],
        "tables": [
            {
                "caption": "Table 1: Dielectric constants of common materials",
                "headers": ["Material", "κ (relative)", "Breakdown (kV/mm)"],
                "rows": [
                    ["Vacuum", "1.0000", "—"],
                    ["Air", "1.0006", "3.0"],
                    ["Paper", "3.7", "16"],
                    ["Glass", "5.6", "14"],
                    ["Water (20°C)", "80.1", "—"],
                    ["Titanium dioxide", "86–173", "4"],
                ],
            },
        ],
        "equations": [
            ("F = kq₁q₂/r²", "Equation 1"),
            ("E = F/q", "Equation 2"),
            ("V = kq/r", "Equation 3"),
        ],
    },
    {
        "title": "Chapter 5: Magnetism",
        "intro": "Magnetism arises from moving electric charges and is intimately connected to electricity through Maxwell's equations. This chapter explores magnetic fields, forces on conductors, and electromagnetic induction.",
        "sections": [
            ("5.1 Magnetic Fields and Forces", "Magnetic fields exert forces on moving charges and current-carrying conductors. The direction of the force is perpendicular to both the velocity and the field, following the right-hand rule."),
            ("5.2 Electromagnetic Induction", "Faraday's law states that a changing magnetic flux through a circuit induces an electromotive force. This principle underlies the operation of generators, transformers, and many other devices."),
        ],
        "figures": [
            ("Figure 1: Magnetic field around a current-carrying solenoid", "Cross-section of a solenoid with N = 500 turns, length 0.3 m, carrying I = 2 A, showing uniform interior field and diverging exterior field lines."),
            ("Figure 2: Lenz's law demonstration with a falling magnet", "Sequence showing a bar magnet falling through a copper ring, with induced current direction and retarding force indicated at each stage."),
            ("Figure 3: Transformer schematic with primary and secondary coils", "An iron-core transformer with N₁ = 1000 turns (primary) and N₂ = 50 turns (secondary) for step-down voltage conversion from 220V to 11V."),
        ],
        "tables": [
            {
                "caption": "Table 1: Magnetic susceptibility of selected materials",
                "headers": ["Material", "χ_m", "Type"],
                "rows": [
                    ["Bismuth", "-1.66 × 10⁻⁴", "Diamagnetic"],
                    ["Copper", "-9.63 × 10⁻⁶", "Diamagnetic"],
                    ["Aluminum", "2.22 × 10⁻⁵", "Paramagnetic"],
                    ["Platinum", "2.63 × 10⁻⁴", "Paramagnetic"],
                    ["Iron", "~200,000", "Ferromagnetic"],
                ],
            },
            {
                "caption": "Table 2: Common electromagnetic devices and principles",
                "headers": ["Device", "Principle", "Application"],
                "rows": [
                    ["Electric motor", "Lorentz force on current loop", "Mechanical rotation"],
                    ["Generator", "Faraday's law", "Electricity production"],
                    ["Transformer", "Mutual inductance", "Voltage conversion"],
                    ["Electromagnet", "Solenoid field", "Material handling"],
                ],
            },
        ],
        "equations": [
            ("F = qv × B", "Equation 1"),
            ("B = μ₀nI", "Equation 2"),
            ("ε = -dΦ_B/dt", "Equation 3"),
            ("V₂/V₁ = N₂/N₁", "Equation 4"),
        ],
    },
    {
        "title": "Chapter 6: Optics",
        "intro": "Optics deals with the behavior and properties of light and its interactions with matter. This chapter covers geometric optics, wave optics, and the nature of electromagnetic radiation.",
        "sections": [
            ("6.1 Reflection and Refraction", "When light encounters a boundary between two media, it can be reflected, refracted, or both. Snell's law quantifies the relationship between angles of incidence and refraction."),
            ("6.2 Interference and Diffraction", "The wave nature of light gives rise to interference and diffraction patterns. Young's double-slit experiment provided key evidence for the wave theory of light."),
        ],
        "figures": [
            ("Figure 1: Ray diagram for a converging lens", "Principal rays through a thin convex lens with focal length f = 15 cm, forming a real inverted image of an object placed at d_o = 25 cm."),
            ("Figure 2: Double-slit interference pattern", "Intensity distribution on a screen 2.0 m from two slits separated by d = 0.25 mm with λ = 550 nm monochromatic light."),
        ],
        "tables": [
            {
                "caption": "Table 1: Refractive indices of common transparent materials",
                "headers": ["Material", "n (at 589 nm)", "Critical Angle (°)"],
                "rows": [
                    ["Air", "1.000", "—"],
                    ["Water", "1.333", "48.8"],
                    ["Crown glass", "1.523", "41.0"],
                    ["Flint glass", "1.660", "37.0"],
                    ["Diamond", "2.417", "24.4"],
                ],
            },
        ],
        "equations": [
            ("n₁ sin θ₁ = n₂ sin θ₂", "Equation 1"),
            ("1/f = 1/d_o + 1/d_i", "Equation 2"),
            ("d sin θ = mλ", "Equation 3"),
        ],
    },
    {
        "title": "Chapter 7: Modern Physics",
        "intro": "Modern physics encompasses quantum mechanics and relativity, two revolutionary frameworks developed in the early 20th century that fundamentally changed our understanding of nature at very small and very fast scales.",
        "sections": [
            ("7.1 Special Relativity", "Einstein's special theory of relativity is based on two postulates: the laws of physics are the same in all inertial frames, and the speed of light in vacuum is constant for all observers."),
            ("7.2 Quantum Mechanics", "Quantum mechanics describes the behavior of matter and energy at atomic and subatomic scales. Key concepts include wave-particle duality, the uncertainty principle, and quantization of energy."),
        ],
        "figures": [
            ("Figure 1: Photoelectric effect apparatus and results", "Schematic of a vacuum phototube connected to a variable voltage source, with a graph of photocurrent vs. voltage for different light frequencies."),
            ("Figure 2: Energy level diagram for hydrogen atom", "Bohr model energy levels from n=1 to n=6, showing the Lyman, Balmer, and Paschen spectral series with wavelengths labeled."),
            ("Figure 3: Compton scattering geometry", "Diagram showing an incident X-ray photon scattering off a stationary electron, with scattered photon and recoiling electron at angles θ and φ respectively."),
        ],
        "tables": [
            {
                "caption": "Table 1: Fundamental particles of the Standard Model",
                "headers": ["Particle", "Charge (e)", "Mass (MeV/c²)", "Spin"],
                "rows": [
                    ["Electron", "-1", "0.511", "1/2"],
                    ["Proton", "+1", "938.3", "1/2"],
                    ["Neutron", "0", "939.6", "1/2"],
                    ["Photon", "0", "0", "1"],
                    ["Muon", "-1", "105.7", "1/2"],
                    ["Pion (π⁺)", "+1", "139.6", "0"],
                ],
            },
        ],
        "equations": [
            ("E = mc²", "Equation 1"),
            ("E = hf", "Equation 2"),
            ("ΔxΔp ≥ ℏ/2", "Equation 3"),
            ("λ = h/p", "Equation 4"),
        ],
    },
    {
        "title": "Chapter 8: Nuclear Physics",
        "intro": "Nuclear physics studies the constituents and interactions of atomic nuclei. The strong nuclear force binds protons and neutrons together, and nuclear reactions release enormous amounts of energy according to Einstein's mass-energy equivalence.",
        "sections": [
            ("8.1 Nuclear Structure and Binding Energy", "Atomic nuclei are composed of protons and neutrons (collectively called nucleons) held together by the strong nuclear force. The binding energy represents the energy required to disassemble a nucleus into its constituent nucleons."),
            ("8.2 Radioactive Decay", "Unstable nuclei undergo radioactive decay, emitting alpha particles, beta particles, or gamma rays. Each radioactive isotope has a characteristic half-life."),
        ],
        "figures": [
            ("Figure 1: Binding energy per nucleon curve", "Plot of B/A vs. mass number A from A=1 to A=240, showing the peak at iron-56 (8.8 MeV/nucleon) and the implications for fission and fusion."),
            ("Figure 2: Decay chain of Uranium-238", "Diagram showing the complete decay series from U-238 through intermediate isotopes to stable Pb-206, with half-lives and decay types labeled."),
        ],
        "tables": [
            {
                "caption": "Table 1: Properties of common radioactive isotopes",
                "headers": ["Isotope", "Half-life", "Decay Mode", "Application"],
                "rows": [
                    ["Carbon-14", "5,730 yr", "β⁻", "Archaeological dating"],
                    ["Cobalt-60", "5.27 yr", "β⁻, γ", "Medical therapy"],
                    ["Technetium-99m", "6.01 hr", "γ", "Medical imaging"],
                    ["Uranium-235", "7.04 × 10⁸ yr", "α", "Nuclear fuel"],
                    ["Plutonium-239", "2.41 × 10⁴ yr", "α", "Nuclear weapons"],
                ],
            },
        ],
        "equations": [
            ("N(t) = N₀ e^(-λt)", "Equation 1"),
            ("t₁/₂ = ln(2)/λ", "Equation 2"),
            ("E = Δmc²", "Equation 3"),
        ],
    },
]


def create_initial():
    doc = Document()

    # Title page
    doc.add_heading("Comprehensive Physics Textbook", level=0)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Master Document — Third Edition")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)

    p2 = doc.add_paragraph()
    p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = p2.add_run("Prof. Elena Vasquez & Prof. James Rutherford")
    run2.font.size = Pt(13)

    p3 = doc.add_paragraph()
    p3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run3 = p3.add_run("Department of Physics, Meridian University\n2025")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Table of Contents placeholder
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    for ch in chapters:
        toc_entry = doc.add_paragraph(ch["title"])
        toc_entry.paragraph_format.left_indent = Inches(0.3)

    # Build each chapter
    for ch in chapters:
        # Section break for each chapter
        new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)

        add_heading_styled(doc, ch["title"], level=1)
        doc.add_paragraph(ch["intro"])

        fig_idx = 0
        eq_idx = 0
        tbl_idx = 0

        # Interleave content: sections, then figures/tables/equations
        for sec_title, sec_text in ch["sections"]:
            add_heading_styled(doc, sec_title, level=2)
            doc.add_paragraph(sec_text)

            # Add a figure after each section if available
            if fig_idx < len(ch["figures"]):
                caption, desc = ch["figures"][fig_idx]
                doc.add_paragraph(f"[{desc}]")
                add_caption(doc, caption)
                fig_idx += 1

            # Add an equation after each section if available
            if eq_idx < len(ch["equations"]):
                eq_text, eq_label = ch["equations"][eq_idx]
                add_equation_block(doc, eq_text, eq_label)
                eq_idx += 1

        # Add remaining figures
        while fig_idx < len(ch["figures"]):
            caption, desc = ch["figures"][fig_idx]
            doc.add_paragraph(f"[{desc}]")
            add_caption(doc, caption)
            fig_idx += 1

        # Add remaining equations
        while eq_idx < len(ch["equations"]):
            eq_text, eq_label = ch["equations"][eq_idx]
            add_equation_block(doc, eq_text, eq_label)
            eq_idx += 1

        # Add tables at end of chapter
        for tbl in ch["tables"]:
            add_data_table(doc, tbl["headers"], tbl["rows"], tbl["caption"])

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
