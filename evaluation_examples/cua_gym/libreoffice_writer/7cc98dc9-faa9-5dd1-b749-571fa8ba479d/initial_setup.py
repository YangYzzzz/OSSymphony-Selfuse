"""
Initial Setup: research_report.docx for writer_edit_034
Task ID: writer_edit_034
Domain: libreoffice_writer

Creates a 10-page research report document with multiple headings.
'Section 3: Methodology' is on page 5. The paragraph immediately after
this heading starts with 'Data was collected from...' (no intro paragraph yet).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_034'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ─── Title Page ───────────────────────────────────────────────────────────
    title_para = doc.add_heading('Understanding Climate Adaptation Strategies\nin Urban Environments', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')
    sub = doc.add_paragraph('A Comprehensive Research Report')
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')
    authors = doc.add_paragraph('Dr. Elena Vasquez · Prof. James Okafor · Dr. Mei-Ling Huang')
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')
    affil = doc.add_paragraph('Department of Environmental Sciences\nGreenfield University')
    affil.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')
    date_para = doc.add_paragraph('March 2025')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # ─── Abstract ─────────────────────────────────────────────────────────────
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This report examines how cities across four continents are adapting to the '
        'accelerating impacts of climate change. Drawing on field surveys, remote '
        'sensing data, and 47 structured interviews with urban planners and local '
        'government representatives, we identify six distinct adaptation archetypes '
        'and map their distribution against socioeconomic vulnerability indices. '
        'Our findings suggest that hybrid green-grey infrastructure programmes reduce '
        'heat-island severity by 1.8–3.4 °C and cut peak stormwater runoff by up to '
        '34 %. Policy recommendations and an implementation roadmap are provided for '
        'municipalities with populations between 250 000 and 2 million.'
    )

    doc.add_page_break()

    # ─── Section 1: Introduction ──────────────────────────────────────────────
    doc.add_heading('Section 1: Introduction', level=1)
    doc.add_paragraph(
        'Urban areas house more than 55 % of the global population and are '
        'responsible for approximately 70 % of global greenhouse-gas emissions. '
        'At the same time, the concentration of people, infrastructure, and economic '
        'activity in cities amplifies their exposure to climate hazards such as '
        'extreme heat, flooding, and sea-level rise. Municipal governments are '
        'therefore under growing pressure to develop and implement coherent climate '
        'adaptation strategies that safeguard residents while maintaining economic '
        'competitiveness.'
    )
    doc.add_paragraph(
        'Despite a decade of policy experimentation, significant knowledge gaps '
        'remain. Most existing literature focuses on single-hazard scenarios or '
        'individual cities, limiting transferability of lessons. This report aims '
        'to fill that gap by synthesising evidence from 24 cities across North '
        'America, Europe, Africa, and South-East Asia, providing a comparative '
        'framework grounded in both quantitative indicators and qualitative '
        'stakeholder perspectives.'
    )

    doc.add_heading('1.1 Background and Motivation', level=2)
    doc.add_paragraph(
        'The Intergovernmental Panel on Climate Change (IPCC) Sixth Assessment '
        'Report (AR6) projects that without deep mitigation, average urban '
        'temperatures will rise by 2–4 °C above pre-industrial levels by 2080, '
        'with coastal cities facing storm-surge events previously classified as '
        '100-year episodes occurring every 10–15 years by mid-century. These '
        'projections have galvanised municipal action, yet a consistent and '
        'replicable adaptation methodology is still absent from the literature.'
    )

    doc.add_heading('1.2 Research Questions', level=2)
    rq_items = [
        'What adaptation strategies have cities deployed most consistently across different climate zones?',
        'Which factors predict the effectiveness of green versus grey infrastructure investment?',
        'How do governance structures influence the speed and coherence of adaptation roll-out?',
        'What socioeconomic disparities exist in access to adaptation benefits within cities?',
    ]
    for item in rq_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ─── Section 2: Literature Review ─────────────────────────────────────────
    doc.add_heading('Section 2: Literature Review', level=1)
    doc.add_paragraph(
        'The academic literature on urban climate adaptation has grown substantially '
        'since the early 2000s. Seminal work by Adger et al. (2005) framed '
        'adaptation as a social-ecological systems problem, emphasising the role of '
        'institutional capacity and community resilience. Subsequent meta-analyses '
        'have confirmed that adaptive capacity is closely linked to governance '
        'quality (Biesbroek et al., 2010; Preston et al., 2011).'
    )
    doc.add_paragraph(
        'Green infrastructure — including urban forests, green roofs, and '
        'constructed wetlands — has attracted particular attention as a multi-benefit '
        'adaptation tool. A systematic review by Faivre et al. (2017) found that '
        'green infrastructure consistently reduced surface temperatures, improved '
        'stormwater management, and enhanced biodiversity in urban settings. '
        'However, the authors cautioned that maintenance costs and land-value '
        'pressures often erode long-term sustainability.'
    )

    doc.add_heading('2.1 Adaptation Typologies', level=2)
    doc.add_paragraph(
        'Several researchers have proposed typological frameworks for classifying '
        'urban adaptation measures. Hallegatte (2009) distinguished between '
        '"no-regret" actions (beneficial under all climate scenarios), "reversible" '
        'actions (easily modified as conditions change), and "transformative" actions '
        '(fundamental changes to urban form or governance). This tripartite '
        'classification remains influential in policy guidance documents.'
    )
    doc.add_paragraph(
        'More recent work by Revi et al. (2014) in the IPCC AR5 introduced the '
        'concept of "adaptation pathways" — sequences of policy options that keep '
        'future choices open while addressing near-term vulnerabilities. The pathway '
        'approach has since been operationalised in national adaptation plans in '
        'Bangladesh, the Netherlands, and New Zealand.'
    )

    doc.add_heading('2.2 Gaps in Current Knowledge', level=2)
    doc.add_paragraph(
        'Despite the growing body of evidence, three gaps persist. First, '
        'comparative studies spanning multiple climate zones remain scarce, making '
        'it difficult to distinguish universal principles from context-specific '
        'findings. Second, most evaluations focus on biophysical outcomes and '
        'neglect distributional equity — who benefits and who bears residual risk. '
        'Third, the interaction between digital governance tools (smart-city '
        'platforms, early-warning systems) and traditional adaptation measures is '
        'poorly understood.'
    )

    doc.add_page_break()

    # ─── Section 3: Methodology ───────────────────────────────────────────────
    doc.add_heading('Section 3: Methodology', level=1)
    # NOTE: The new paragraph 'This section describes the research methodology
    # used in this study.' is NOT present in the initial state — that is the task.
    doc.add_paragraph(
        'Data was collected from 24 cities through a mixed-methods design combining '
        'quantitative analysis of municipal administrative records with qualitative '
        'semi-structured interviews. Cities were selected using a stratified sampling '
        'frame based on climate zone (Köppen classification), population size, and '
        'GDP per capita, ensuring representation across high-, middle-, and '
        'low-income contexts.'
    )
    doc.add_paragraph(
        'Quantitative data included annual temperature anomalies, precipitation '
        'records, and infrastructure expenditure datasets obtained from national '
        'statistical agencies and the World Bank Open Data platform. All monetary '
        'values were converted to 2020 USD using purchasing power parity (PPP) '
        'exchange rates to facilitate cross-country comparison.'
    )

    doc.add_heading('3.1 Sampling Strategy', level=2)
    doc.add_paragraph(
        'The 24 study cities were distributed as follows: 6 in North America '
        '(including Phoenix, Toronto, and Mexico City), 7 in Europe (including '
        'Rotterdam, Copenhagen, and Lisbon), 5 in sub-Saharan Africa (including '
        'Nairobi, Accra, and Durban), and 6 in South and South-East Asia '
        '(including Dhaka, Jakarta, and Ho Chi Minh City). This distribution '
        'was intentional, providing variation in climate hazard profiles, '
        'institutional capacities, and economic resources.'
    )

    doc.add_heading('3.2 Interview Protocol', level=2)
    doc.add_paragraph(
        'Semi-structured interviews were conducted with urban planners, resilience '
        'officers, and community representatives between January and September 2024. '
        'A total of 47 interviews were completed, averaging 72 minutes in length. '
        'Interviews were recorded with participant consent, transcribed verbatim, '
        'and subjected to thematic analysis using NVivo 14. Inter-coder reliability '
        'was established via Cohen\'s kappa (κ = 0.81), indicating strong agreement.'
    )

    doc.add_heading('3.3 Data Quality and Limitations', level=2)
    doc.add_paragraph(
        'Several limitations should be noted. Administrative data quality varied '
        'markedly across cities; in three African cases, temperature records had '
        'gaps exceeding 18 months, necessitating interpolation. Self-selection '
        'bias may affect interview findings, as officials willing to participate '
        'may be more engaged with adaptation agendas than average. Finally, the '
        'cross-sectional design cannot establish causality between governance '
        'structures and adaptation outcomes.'
    )

    doc.add_page_break()

    # ─── Section 4: Findings ──────────────────────────────────────────────────
    doc.add_heading('Section 4: Findings', level=1)
    doc.add_paragraph(
        'Analysis of the combined quantitative and qualitative dataset reveals '
        'six recurring adaptation archetypes, which we label: (1) Green-Grey '
        'Hybrid, (2) Nature-Based Pioneer, (3) Technology-Led Resilience, '
        '(4) Community-Driven Incremental, (5) Regulatory Compliance, and '
        '(6) Laggard-Under-Constraint. Each archetype is described below with '
        'representative city examples.'
    )

    doc.add_heading('4.1 Green-Grey Hybrid Cities', level=2)
    doc.add_paragraph(
        'Fourteen of the 24 cities exhibited features of the Green-Grey Hybrid '
        'archetype, combining engineered flood defences (grey infrastructure) with '
        'urban greening programmes. Rotterdam exemplifies this model through its '
        '"room for the river" policy, integrating tidal barriers with water plazas '
        'and floating parks. Surface temperature monitoring showed a 2.9 °C '
        'reduction in green-grey corridor zones compared with comparable '
        'unmodified districts.'
    )

    doc.add_heading('4.2 Nature-Based Pioneer Cities', level=2)
    doc.add_paragraph(
        'Four cities — Copenhagen, Nairobi, Ho Chi Minh City, and Singapore — '
        'were classified as Nature-Based Pioneers due to their early adoption of '
        'ecosystem-based adaptation at scale. Copenhagen\'s climate-resilient '
        'neighbourhood plan (Klimatilpasningsplan) redirected stormwater into '
        'channelled urban streams and recreational parks, reducing combined-sewer '
        'overflow incidents by 62 % between 2015 and 2023.'
    )

    doc.add_page_break()

    # ─── Section 5: Discussion ────────────────────────────────────────────────
    doc.add_heading('Section 5: Discussion', level=1)
    doc.add_paragraph(
        'The typological analysis supports the assertion that adaptation is not a '
        'single technical solution but a context-sensitive portfolio of measures '
        'shaped by institutional capacity, financial resources, and political will. '
        'Cities that demonstrated the strongest adaptation performance shared three '
        'common features: long-term vision embedded in statutory plans, dedicated '
        'adaptation budgets insulated from electoral cycles, and active community '
        'engagement mechanisms.'
    )
    doc.add_paragraph(
        'Equity considerations deserve particular emphasis. In seven of the eight '
        'cities where socioeconomic disaggregation was possible, adaptation '
        'benefits disproportionately accrued to higher-income districts. This '
        'finding echoes Shi et al. (2016) and underscores the need for explicitly '
        'equity-oriented adaptation planning frameworks.'
    )

    doc.add_heading('5.1 Policy Implications', level=2)
    policy_items = [
        'Establish ring-fenced climate adaptation funds at municipal level.',
        'Mandate equity impact assessments for all major adaptation investments.',
        'Create cross-departmental resilience offices with statutory authority.',
        'Develop standardised monitoring frameworks to enable longitudinal comparison.',
        'Invest in capacity building for community-based organisations in high-risk '
        'neighbourhoods.',
    ]
    for item in policy_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ─── Section 6: Conclusions ────────────────────────────────────────────────
    doc.add_heading('Section 6: Conclusions', level=1)
    doc.add_paragraph(
        'This report has demonstrated that effective urban climate adaptation '
        'requires the simultaneous mobilisation of technical, institutional, and '
        'social resources. No single archetype dominates across all contexts; '
        'rather, cities must select and sequence measures according to their '
        'specific hazard profiles, governance structures, and financial capacity. '
        'The adaptation pathway framework offers a useful conceptual tool for '
        'structuring this sequencing process.'
    )
    doc.add_paragraph(
        'Future research should prioritise longitudinal evaluation designs to '
        'establish causal relationships between adaptation investments and '
        'resilience outcomes. Additionally, improved methods for capturing '
        'distributional effects are needed to ensure that adaptation policy '
        'reduces rather than reinforces existing inequalities.'
    )

    doc.add_page_break()

    # ─── References ───────────────────────────────────────────────────────────
    doc.add_heading('References', level=1)
    references = [
        'Adger, W. N., Arnell, N. W., & Tompkins, E. L. (2005). Successful adaptation to climate change across scales. Global Environmental Change, 15(2), 77–86.',
        'Biesbroek, G. R., Swart, R. J., Carter, T. R., Cowan, C., Henrichs, T., Mela, H., Morecroft, M. D., & Rey, D. (2010). Europe adapts to climate change: comparing national adaptation strategies. Global Environmental Change, 20(3), 440–450.',
        'Faivre, N., Fritz, M., Freitas, T., de Boissezon, B., & Vandewoestijne, S. (2017). Nature-based solutions in the EU: innovating with nature to address social, economic and environmental challenges. Environmental Research, 159, 509–518.',
        'Hallegatte, S. (2009). Strategies to adapt to an uncertain climate change. Global Environmental Change, 19(2), 240–247.',
        'IPCC (2022). Sixth Assessment Report: Impacts, Adaptation and Vulnerability. Cambridge University Press.',
        'Preston, B. L., Westaway, R. M., & Yuen, E. J. (2011). Climate adaptation planning in practice: an evaluation of adaptation plans from three developed nations. Mitigation and Adaptation Strategies for Global Change, 16(4), 407–438.',
        'Revi, A., Satterthwaite, D., Aragón-Durand, F., Corfee-Morlot, J., Kiunsi, R. B. R., Pelling, M., Roberts, D., Solecki, W., Gajjar, S. P., & Sverdlik, A. (2014). Towards transformative adaptation in cities: the IPCC\'s Fifth Assessment. Environment and Urbanization, 26(1), 11–28.',
        'Shi, L., Chu, E., Anguelovski, I., Aylett, A., Debats, J., Goh, K., Schenk, T., Seto, K. C., Dodman, D., Roberts, D., Roberts, J. T., & VanDeveer, S. D. (2016). Roadmap towards justice in urban climate adaptation research. Nature Climate Change, 6(2), 131–137.',
    ]
    for ref in references:
        p = doc.add_paragraph(ref, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Copy to Desktop as per context description
    desktop_path = f'{WORKDIR}/Desktop/research_report.docx'
    import shutil
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)
    shutil.copy(OUTPUT, desktop_path)
    print(f'Copied to Desktop: {desktop_path}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
