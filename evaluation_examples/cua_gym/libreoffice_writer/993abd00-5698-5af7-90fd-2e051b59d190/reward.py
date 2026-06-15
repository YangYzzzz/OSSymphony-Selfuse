"""
FINAL REWARD SCRIPT - SUCCESS
Task: For the first two paragraphs, bump the line spacing up to 2×.
Generated: 2025-10-14 09:05:04
Status: success
Model: azure-o3
Total Steps: 2
"""

from docx import Document
from docx.enum.text import WD_LINE_SPACING
import os


def verify_double_spacing_first_two(file_path: str) -> float:
    """Verify that the first two paragraphs of a DOCX file are set to 2× (double) line spacing.

    Scoring (progressive):
        • 0.5 points for each of the first two paragraphs that is double-spaced
        • Maximum possible score = 1.0

    The function prints diagnostic information and the final reward ("REWARD: X.X").
    """

    print(f"Verifying document: {file_path}")

    # ---------- Preliminary checks ----------
    if not os.path.exists(file_path):
        print(f"✗ File not found: {file_path}")
        print("REWARD: 0.0")
        return 0.0  # Cannot proceed without the file

    try:
        doc = Document(file_path)
        print(f"✓ Document loaded successfully with {len(doc.paragraphs)} paragraph(s)")
    except Exception as e:
        print(f"✗ Error loading document: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Must have at least two paragraphs to satisfy the task
    if len(doc.paragraphs) < 2:
        print("✗ Document contains fewer than 2 paragraphs. Task cannot be fulfilled.")
        print("REWARD: 0.0")
        return 0.0

    # ---------- Core verification ----------
    double_spaced_count = 0
    paragraphs_to_check = 2  # Only the first two paragraphs matter

    for idx in range(paragraphs_to_check):
        para = doc.paragraphs[idx]
        pf = para.paragraph_format

        # Extract spacing information
        line_spacing_val = pf.line_spacing         # Float, Length, or None
        line_spacing_rule = pf.line_spacing_rule   # Enum SINGLE, DOUBLE, etc.

        is_double = False

        # Method 1: Explicit DOUBLE rule
        if line_spacing_rule == WD_LINE_SPACING.DOUBLE:
            is_double = True
        else:
            # Method 2: Multiple spacing expressed as a float (≈ 2.0)
            if isinstance(line_spacing_val, (int, float)) and line_spacing_val is not None:
                if abs(float(line_spacing_val) - 2.0) < 0.05:
                    is_double = True

        print(
            f"Paragraph {idx + 1}: line_spacing_rule={line_spacing_rule}, "
            f"line_spacing_val={line_spacing_val} -> double={is_double}"
        )

        if is_double:
            double_spaced_count += 1

    # ---------- Scoring ----------
    total_score = 0.5 * double_spaced_count  # 0.5 per correctly formatted paragraph
    final_score = min(total_score, 1.0)      # Cap at 1.0

    print(f"✓ Double-spaced paragraphs: {double_spaced_count}/{paragraphs_to_check}")
    print(f"REWARD: {final_score}")

    return final_score


# -----------------------------
# MAIN EXECUTION (Required)
# -----------------------------
if __name__ == "__main__":
    FILE_PATH = "/home/user/for_the_first_two_paragraphs_bump_the_line_spacing_up_to_2.docx"
    verify_double_spacing_first_two(FILE_PATH)

