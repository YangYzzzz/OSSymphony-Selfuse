"""
Initial Setup: Tradeshow handout document in single-column portrait layout
Task ID: writer_mktg_035
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_035'
OUTPUT = f'{WORKDIR}/Desktop/tradeshow_handout.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # --- Page setup: portrait, letter, 1-inch margins, single column ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Normal (no column definition — single column by default)
    # No orientation override — portrait is default

    # --- Content: company name ---
    p = doc.add_paragraph()
    run = p.add_run('Apex Dynamics')
    run.font.size = Pt(12)

    # --- Tagline ---
    p2 = doc.add_paragraph()
    run2 = p2.add_run('Intelligent Automation for Modern Enterprise')
    run2.font.size = Pt(12)

    # --- Product highlights (single paragraph) ---
    doc.add_paragraph()  # blank line separator
    p3 = doc.add_paragraph()
    run3 = p3.add_run(
        'Product Highlights: '
        'SmartControl Platform — AI-powered process automation with real-time monitoring. '
        'DataBridge Suite — Seamless integration across legacy and cloud systems. '
        'PredictIQ Engine — Machine learning forecasting for operational efficiency. '
        'SecureOps Shield — Zero-trust cybersecurity built for enterprise compliance.'
    )
    run3.font.size = Pt(12)

    # --- Contact information ---
    doc.add_paragraph()  # blank line separator
    p4 = doc.add_paragraph()
    run4 = p4.add_run(
        'Contact Us: '
        '1200 Innovation Drive, Suite 400, Austin, TX 78701. '
        'Phone: (512) 555-0198. '
        'Email: info@apexdynamics.com. '
        'Website: www.apexdynamics.com.'
    )
    run4.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
