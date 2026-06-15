"""
Initial Setup: Remove all list formatting from the entire document
Task ID: writer_lec_032
Domain: libreoffice_writer

Creates a Writer document with mixed content: 3 bulleted lists and 2 numbered lists
scattered throughout ~4 pages of text. All text is realistic business content.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_032'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page 1: Introduction and first bulleted list ---
    h = doc.add_heading("Quarterly Business Review - Q1 2025", level=1)

    doc.add_paragraph(
        "This report provides a comprehensive overview of the company's performance "
        "during the first quarter of 2025. The following sections cover financial "
        "highlights, operational milestones, and strategic initiatives that shaped "
        "our results."
    )

    doc.add_paragraph(
        "The leadership team has identified several key areas of growth that "
        "contributed to a strong start to the fiscal year. Revenue exceeded "
        "expectations by 12%, driven primarily by expansion in the Asia-Pacific "
        "region and a successful product launch in February."
    )

    doc.add_heading("Financial Highlights", level=2)

    doc.add_paragraph(
        "Our financial performance this quarter reflects the culmination of "
        "strategic investments made throughout 2024. Below are the primary "
        "revenue drivers identified by the finance team:"
    )

    # Bulleted list 1: Revenue drivers
    doc.add_paragraph("Enterprise software licensing grew 18% year-over-year, reaching $42.3 million", style="List Bullet")
    doc.add_paragraph("Cloud services subscription revenue increased to $28.7 million, up from $21.4 million in Q1 2024", style="List Bullet")
    doc.add_paragraph("Professional services engagements contributed $15.2 million, a 9% improvement", style="List Bullet")
    doc.add_paragraph("Hardware maintenance contracts generated $8.9 million in recurring revenue", style="List Bullet")
    doc.add_paragraph("New market penetration in Southeast Asia added $6.1 million in first-time revenue", style="List Bullet")

    doc.add_paragraph(
        "Total revenue for the quarter reached $101.2 million, representing a "
        "combined growth rate of 14.3% compared to the same period last year. "
        "Gross margins improved by 2.1 percentage points to 67.8%, reflecting "
        "improved operational efficiency and favorable product mix."
    )

    doc.add_paragraph(
        "Operating expenses were managed within budget at $52.4 million. Research "
        "and development spending accounted for 22% of revenue, consistent with "
        "our long-term innovation strategy. Sales and marketing costs decreased "
        "as a percentage of revenue due to improved lead conversion rates."
    )

    # --- Page 2: Operational milestones with numbered list ---
    doc.add_heading("Operational Milestones", level=2)

    doc.add_paragraph(
        "The operations team achieved several significant milestones during Q1 "
        "that position the company for continued growth. These achievements "
        "span product development, infrastructure modernization, and talent "
        "acquisition. Each milestone was tracked against our annual OKR framework."
    )

    # Numbered list 1: Milestones
    doc.add_paragraph("Launched the Aurora 3.0 platform with advanced analytics capabilities on January 15, completing a 14-month development cycle", style="List Number")
    doc.add_paragraph("Migrated 87% of legacy infrastructure to cloud-native architecture, reducing hosting costs by $1.2 million annually", style="List Number")
    doc.add_paragraph("Expanded the engineering team by 34 new hires across three offices in San Francisco, Berlin, and Singapore", style="List Number")
    doc.add_paragraph("Achieved SOC 2 Type II certification ahead of schedule, enabling entry into regulated industry verticals", style="List Number")
    doc.add_paragraph("Established a strategic partnership with Meridian Technologies for joint go-to-market in the healthcare sector", style="List Number")
    doc.add_paragraph("Reduced average customer onboarding time from 45 days to 18 days through process automation", style="List Number")

    doc.add_paragraph(
        "The Aurora 3.0 launch was particularly noteworthy, as it represented "
        "the largest product release in the company's history. Customer feedback "
        "has been overwhelmingly positive, with a Net Promoter Score of 72 among "
        "early adopters. The platform's new machine learning features have been "
        "cited as a key differentiator in competitive evaluations."
    )

    doc.add_paragraph(
        "Infrastructure modernization efforts have yielded measurable benefits "
        "beyond cost savings. System uptime improved to 99.97%, and average "
        "page load times decreased by 340 milliseconds. These improvements "
        "directly correlate with increased user engagement metrics across "
        "all product lines."
    )

    # --- Page 3: Strategic initiatives with bulleted list and numbered list ---
    doc.add_heading("Strategic Initiatives", level=2)

    doc.add_paragraph(
        "Looking ahead to Q2 and the remainder of 2025, the executive team "
        "has outlined several strategic priorities. These initiatives are "
        "designed to accelerate growth while maintaining the operational "
        "discipline that characterized our Q1 performance."
    )

    doc.add_paragraph(
        "The board of directors approved a $25 million investment package "
        "to support the following strategic initiatives during their March "
        "meeting. Each initiative has a designated executive sponsor and "
        "quarterly review cadence."
    )

    # Bulleted list 2: Strategic priorities
    doc.add_paragraph("Expand the Aurora platform's AI capabilities with natural language processing and predictive modeling features", style="List Bullet")
    doc.add_paragraph("Enter the Latin American market through a combination of direct sales and channel partnerships with regional distributors", style="List Bullet")
    doc.add_paragraph("Develop an integrated developer ecosystem with APIs, SDKs, and a marketplace for third-party extensions", style="List Bullet")
    doc.add_paragraph("Implement a company-wide sustainability program targeting carbon neutrality by 2027", style="List Bullet")

    doc.add_paragraph(
        "In parallel with these growth initiatives, the company will continue "
        "to invest in customer success programs. Our retention rate of 94.2% "
        "is among the highest in the industry, and we aim to increase it to "
        "96% by year-end through enhanced support offerings and proactive "
        "account management."
    )

    doc.add_heading("Risk Factors and Mitigation", level=2)

    doc.add_paragraph(
        "While the outlook for 2025 is positive, the leadership team has "
        "identified several risk factors that warrant monitoring. The "
        "competitive landscape continues to evolve, with two well-funded "
        "startups entering our core market segment in February."
    )

    # Numbered list 2: Risk factors
    doc.add_paragraph("Increased competition in the enterprise analytics space from both established players and venture-backed startups", style="List Number")
    doc.add_paragraph("Potential supply chain disruptions affecting hardware delivery timelines for on-premise installations", style="List Number")
    doc.add_paragraph("Regulatory changes in the European Union regarding data processing requirements under the revised Digital Services Act", style="List Number")
    doc.add_paragraph("Currency fluctuation exposure due to expanded international operations, particularly in emerging markets", style="List Number")
    doc.add_paragraph("Talent retention challenges in key engineering roles given competitive hiring market conditions", style="List Number")

    doc.add_paragraph(
        "To mitigate these risks, the company has established a dedicated "
        "competitive intelligence team, diversified its supply chain across "
        "three continents, and engaged specialized legal counsel for EU "
        "regulatory compliance. A currency hedging strategy has been "
        "implemented covering 80% of projected international revenue."
    )

    # --- Page 4: Team updates with bulleted list ---
    doc.add_heading("Team and Culture", level=2)

    doc.add_paragraph(
        "Our people remain our greatest asset. The human resources team "
        "completed several important initiatives during Q1 that strengthen "
        "our organizational capabilities and workplace culture."
    )

    doc.add_paragraph(
        "Employee satisfaction scores from the annual survey reached an "
        "all-time high of 4.3 out of 5.0, with particular improvements "
        "in the categories of career development and work-life balance. "
        "The following programs contributed to these results:"
    )

    # Bulleted list 3: HR programs
    doc.add_paragraph("Launched a mentorship program pairing 120 junior employees with senior leaders across all departments", style="List Bullet")
    doc.add_paragraph("Introduced flexible work arrangements allowing employees to choose between remote, hybrid, and in-office schedules", style="List Bullet")
    doc.add_paragraph("Expanded the learning and development budget by 30%, providing access to external certifications and conference attendance", style="List Bullet")
    doc.add_paragraph("Created four new employee resource groups focused on diversity, wellness, sustainability, and innovation", style="List Bullet")
    doc.add_paragraph("Rolled out an enhanced parental leave policy offering 16 weeks of fully paid leave for all new parents", style="List Bullet")

    doc.add_paragraph(
        "As we move into Q2, the company is well-positioned to build on the "
        "momentum established in the first quarter. The combination of strong "
        "financial results, operational achievements, and strategic investments "
        "provides a solid foundation for achieving our full-year targets."
    )

    doc.add_heading("Conclusion", level=2)

    doc.add_paragraph(
        "The first quarter of 2025 demonstrated that our strategic vision is "
        "translating into measurable results. We remain committed to delivering "
        "value for our customers, employees, and shareholders while pursuing "
        "sustainable long-term growth. The leadership team is confident that "
        "the initiatives outlined in this review will drive continued success "
        "throughout the remainder of the fiscal year."
    )

    doc.add_paragraph(
        "This report was prepared by the Office of the Chief Financial Officer "
        "in collaboration with department heads across the organization. For "
        "questions or additional detail, please contact the investor relations "
        "team at ir@acmecorp.com or the corporate communications office."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
