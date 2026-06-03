"""
Initial Setup: Document with 5 tracked changes and 3 comments
Task ID: osworld_writer_comment_track_changes_003
Domain: libreoffice_writer

Creates a reviewed contract document containing:
  - 5 tracked changes (mix of insertions and deletions)
  - 3 comment annotations
The agent must reject all tracked changes and remove all comments.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import zipfile
import shutil
import copy

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_comment_track_changes_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# XML namespaces
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14_NS = 'http://schemas.microsoft.com/office/word/2010/wordml'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_w(tag):
    return f'{{{W_NS}}}{tag}'


def make_rpr(author=None, date=None, rev_id=None, rpr_extra=None):
    """Build a w:rPr element, optionally with rPrChange."""
    rpr = OxmlElement('w:rPr')
    if rpr_extra:
        for child in rpr_extra:
            rpr.append(copy.deepcopy(child))
    return rpr


def make_ins_element(rev_id, author, date, run_text, bold=False):
    """Create a tracked insertion: <w:ins><w:r><w:t>text</w:t></w:r></w:ins>"""
    ins = OxmlElement('w:ins')
    ins.set(make_w('id'), str(rev_id))
    ins.set(make_w('author'), author)
    ins.set(make_w('date'), date)

    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b')
        rpr.append(b)
    r.append(rpr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = run_text
    r.append(t)
    ins.append(r)
    return ins


def make_del_element(rev_id, author, date, run_text):
    """Create a tracked deletion: <w:del><w:r><w:delText>text</w:delText></w:r></w:del>"""
    del_elem = OxmlElement('w:del')
    del_elem.set(make_w('id'), str(rev_id))
    del_elem.set(make_w('author'), author)
    del_elem.set(make_w('date'), date)

    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    r.append(rpr)
    dt = OxmlElement('w:delText')
    dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    dt.text = run_text
    r.append(dt)
    del_elem.append(r)
    return del_elem


def make_comment_ref(comment_id):
    """Create a w:commentReference element inside a run."""
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(make_w('val'), 'CommentReference')
    rpr.append(rStyle)
    r.append(rpr)
    ref = OxmlElement('w:commentReference')
    ref.set(make_w('id'), str(comment_id))
    r.append(ref)
    return r


def make_comment_range_start(comment_id):
    el = OxmlElement('w:commentRangeStart')
    el.set(make_w('id'), str(comment_id))
    return el


def make_comment_range_end(comment_id):
    el = OxmlElement('w:commentRangeEnd')
    el.set(make_w('id'), str(comment_id))
    return el


def make_normal_run(text, bold=False):
    """Create a normal (untracked) run."""
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b')
        rpr.append(b)
    r.append(rpr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    return r


def build_comments_xml(comments_data):
    """Build the comments.xml content string."""
    root = etree.Element(
        '{%s}comments' % W_NS,
        nsmap={
            'wpc': 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas',
            'w': W_NS,
            'w14': W14_NS,
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        }
    )
    for cid, author, date, text in comments_data:
        comment = etree.SubElement(root, '{%s}comment' % W_NS)
        comment.set('{%s}id' % W_NS, str(cid))
        comment.set('{%s}author' % W_NS, author)
        comment.set('{%s}date' % W_NS, date)
        comment.set('{%s}initials' % W_NS, author.split()[0][0] + (author.split()[1][0] if len(author.split()) > 1 else ''))
        p = etree.SubElement(comment, '{%s}p' % W_NS)
        ppr = etree.SubElement(p, '{%s}pPr' % W_NS)
        pstyle = etree.SubElement(ppr, '{%s}pStyle' % W_NS)
        pstyle.set('{%s}val' % W_NS, 'CommentText')
        r = etree.SubElement(p, '{%s}r' % W_NS)
        rpr = etree.SubElement(r, '{%s}rPr' % W_NS)
        rStyle = etree.SubElement(rpr, '{%s}rStyle' % W_NS)
        rStyle.set('{%s}val' % W_NS, 'CommentReference')
        ref = etree.SubElement(r, '{%s}annotationRef' % W_NS)
        r2 = etree.SubElement(p, '{%s}r' % W_NS)
        rpr2 = etree.SubElement(r2, '{%s}rPr' % W_NS)
        rStyle2 = etree.SubElement(rpr2, '{%s}rStyle' % W_NS)
        rStyle2.set('{%s}val' % W_NS, 'CommentText')
        t = etree.SubElement(r2, '{%s}t' % W_NS)
        t.text = text
    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


def create_initial():
    doc = Document()

    # ---- Contract Title ----
    title_para = doc.add_paragraph()
    title_para.style = doc.styles['Heading 1']
    title_para.clear()
    title_para._element.append(make_normal_run('SERVICE AGREEMENT', bold=True))

    # ---- Preamble ----
    preamble = doc.add_paragraph()
    preamble.clear()
    preamble._element.append(make_normal_run(
        'This Service Agreement ("Agreement") is entered into as of March 1, 2025, between '
        'Meridian Consulting Group LLC ("Service Provider") and Hartwell Industries Inc. ("Client").'
    ))

    # ---- Section 1: Scope of Services ----
    doc.add_heading('1. Scope of Services', level=2)

    # Para with Comment 0 anchored
    scope_para = doc.add_paragraph()
    scope_para.clear()
    p_el = scope_para._element
    # Comment range for comment 0
    p_el.append(make_comment_range_start(0))
    p_el.append(make_normal_run('The Service Provider agrees to deliver '))
    # Tracked change 1: deletion of "monthly" (rejected → keep "monthly")
    p_el.append(make_del_element(1, 'Rebecca Walsh', '2025-02-14T10:22:00Z', 'monthly'))
    # Tracked change 2: insertion of "quarterly" (rejected → remove "quarterly")
    p_el.append(make_ins_element(2, 'Rebecca Walsh', '2025-02-14T10:22:00Z', 'quarterly'))
    p_el.append(make_normal_run(' performance reports to the Client as outlined in Exhibit A.'))
    p_el.append(make_comment_range_end(0))
    p_el.append(make_comment_ref(0))

    # ---- Section 2: Payment Terms ----
    doc.add_heading('2. Payment Terms', level=2)

    # Para with Comment 1 anchored
    payment_para = doc.add_paragraph()
    payment_para.clear()
    p2_el = payment_para._element
    p2_el.append(make_comment_range_start(1))
    p2_el.append(make_normal_run('The Client shall remit payment of '))
    # Tracked change 3: deletion of "$12,500" (rejected → keep "$12,500")
    p2_el.append(make_del_element(3, 'James Thornton', '2025-02-18T14:05:00Z', '$12,500'))
    # Tracked change 4: insertion of "$15,000" (rejected → remove "$15,000")
    p2_el.append(make_ins_element(4, 'James Thornton', '2025-02-18T14:05:00Z', '$15,000'))
    p2_el.append(make_normal_run(' per month, due within thirty (30) days of invoice receipt.'))
    p2_el.append(make_comment_range_end(1))
    p2_el.append(make_comment_ref(1))

    # ---- Section 3: Term and Termination ----
    doc.add_heading('3. Term and Termination', level=2)

    term_para = doc.add_paragraph()
    term_para.clear()
    t_el = term_para._element
    t_el.append(make_normal_run('This Agreement shall commence on March 1, 2025 and continue for a period of '))
    # Tracked change 5: insertion of "twelve (12)" replacing nothing (rejected → remove)
    t_el.append(make_ins_element(5, 'Rebecca Walsh', '2025-02-20T09:15:00Z', 'twelve (12) '))
    t_el.append(make_normal_run('months, unless earlier terminated in accordance with this section.'))

    # Para with Comment 2
    termination_para = doc.add_paragraph()
    termination_para.clear()
    te_el = termination_para._element
    te_el.append(make_comment_range_start(2))
    te_el.append(make_normal_run(
        'Either party may terminate this Agreement upon thirty (30) days written notice to the other party.'
    ))
    te_el.append(make_comment_range_end(2))
    te_el.append(make_comment_ref(2))

    # ---- Section 4: Confidentiality ----
    doc.add_heading('4. Confidentiality', level=2)

    conf_para = doc.add_paragraph()
    conf_para.clear()
    conf_para._element.append(make_normal_run(
        'Both parties agree to maintain strict confidentiality regarding all proprietary information, '
        'trade secrets, and business data disclosed during the course of this Agreement.'
    ))

    # ---- Section 5: Governing Law ----
    doc.add_heading('5. Governing Law', level=2)

    law_para = doc.add_paragraph()
    law_para.clear()
    law_para._element.append(make_normal_run(
        'This Agreement shall be governed by the laws of the State of Delaware, '
        'without regard to its conflict of law provisions.'
    ))

    # ---- Signatures ----
    doc.add_heading('Signatures', level=2)
    doc.add_paragraph('Service Provider: _______________________________  Date: ____________')
    doc.add_paragraph('Client: _______________________________  Date: ____________')

    # Save initial document
    doc.save(OUTPUT)

    # Now inject the comments part into the docx using zipfile manipulation
    # python-docx doesn't support comments natively, so we patch the zip

    comments_data = [
        (0, 'Rebecca Walsh', '2025-02-14T10:22:00Z',
         'Consider keeping monthly reports for better tracking of deliverables.'),
        (1, 'James Thornton', '2025-02-18T14:05:00Z',
         'The rate increase should be discussed further with the client before finalizing.'),
        (2, 'Sarah Chen', '2025-02-20T11:30:00Z',
         'Standard 30-day termination clause — acceptable per legal review.'),
    ]
    comments_xml_bytes = build_comments_xml(comments_data)

    # Patch the docx to add comments.xml
    import tempfile
    tmp_path = OUTPUT + '.tmp'
    with zipfile.ZipFile(OUTPUT, 'r') as zin:
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == 'word/_rels/document.xml.rels':
                    # Add relationship for comments
                    rels_content = zin.read(item.filename).decode('utf-8')
                    comment_rel = (
                        '<Relationship Id="rIdComments" '
                        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
                        'Target="comments.xml"/>'
                    )
                    rels_content = rels_content.replace('</Relationships>', comment_rel + '</Relationships>')
                    zout.writestr(item, rels_content.encode('utf-8'))
                elif item.filename == '[Content_Types].xml':
                    # Add content type for comments
                    ct_content = zin.read(item.filename).decode('utf-8')
                    comment_ct = (
                        '<Override PartName="/word/comments.xml" '
                        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
                    )
                    ct_content = ct_content.replace('</Types>', comment_ct + '</Types>')
                    zout.writestr(item, ct_content.encode('utf-8'))
                else:
                    zout.writestr(item, zin.read(item.filename))
            # Write the comments.xml
            zout.writestr('word/comments.xml', comments_xml_bytes)

    shutil.move(tmp_path, OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
