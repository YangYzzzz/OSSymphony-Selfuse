"""
Reward Script: Thesis front matter in LibreOffice Writer
Task ID: writer_wf_032
Domain: libreoffice_writer
Scoring:
  Component 1: Title page content (thesis title, author, university) — 0.25
  Component 2: Abstract section with substantial text — 0.20
  Component 3: Acknowledgments section — 0.15
  Component 4: Table of Contents with 5 chapters — 0.20
  Component 5: List of Tables and List of Figures headings — 0.10
  Component 6: Roman numeral page numbering — 0.10
"""

import os
import re
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_032'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraph texts for searching
    all_para_texts = [p.text.strip() for p in doc.paragraphs]
    all_para_texts_lower = [t.lower() for t in all_para_texts]

    # -----------------------------------------------------------------
    # Component 1: Title page content (0.25 points)
    # The title page must contain the thesis title, author name, and university.
    # These are task-introduced — the initial file is blank.
    # -----------------------------------------------------------------
    try:
        full_text_lower = ' '.join(all_para_texts_lower)
        has_title = 'machine learning applications in predictive maintenance' in full_text_lower
        has_author = 'maria santos' in full_text_lower
        has_university = 'state technical university' in full_text_lower

        comp1 = 0.0
        if has_title:
            comp1 += 0.10
            print("PASS: Thesis title found")
        else:
            print("FAIL: Thesis title 'Machine Learning Applications in Predictive Maintenance' not found")
        if has_author:
            comp1 += 0.08
            print("PASS: Author 'Maria Santos' found")
        else:
            print("FAIL: Author 'Maria Santos' not found")
        if has_university:
            comp1 += 0.07
            print("PASS: University 'State Technical University' found")
        else:
            print("FAIL: University 'State Technical University' not found")

        if comp1 > 0:
            print(f"PASS: Component 1 — Title page content ({comp1:.2f} pts)")
            total_score += comp1
        else:
            print("FAIL: Component 1 — No title page content found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------
    # Component 2: Abstract section with substantial text (0.20 points)
    # Must have "Abstract" heading and a paragraph of significant length.
    # -----------------------------------------------------------------
    try:
        abstract_idx = None
        for i, t in enumerate(all_para_texts_lower):
            if t == 'abstract':
                abstract_idx = i
                break

        if abstract_idx is not None:
            # Look for a substantial paragraph after "Abstract"
            abstract_text = ''
            for j in range(abstract_idx + 1, min(abstract_idx + 5, len(all_para_texts))):
                if all_para_texts[j]:
                    abstract_text = all_para_texts[j]
                    break

            word_count = len(abstract_text.split()) if abstract_text else 0
            if word_count >= 100:
                print(f"PASS: Component 2 — Abstract section found with {word_count} words ({0.20} pts)")
                total_score += 0.20
            elif word_count >= 30:
                partial = 0.10
                print(f"PARTIAL: Component 2 — Abstract exists but only {word_count} words ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — Abstract heading found but text is too short ({word_count} words)")
        else:
            print("FAIL: Component 2 — 'Abstract' heading not found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------
    # Component 3: Acknowledgments section (0.15 points)
    # Must have "Acknowledgments" heading and some text after it.
    # -----------------------------------------------------------------
    try:
        ack_idx = None
        for i, t in enumerate(all_para_texts_lower):
            if t in ('acknowledgments', 'acknowledgements'):
                ack_idx = i
                break

        if ack_idx is not None:
            # Check for text after heading
            ack_text = ''
            for j in range(ack_idx + 1, min(ack_idx + 5, len(all_para_texts))):
                if all_para_texts[j]:
                    ack_text = all_para_texts[j]
                    break

            if len(ack_text.split()) >= 10:
                print(f"PASS: Component 3 — Acknowledgments section found with text ({0.15} pts)")
                total_score += 0.15
            else:
                partial = 0.07
                print(f"PARTIAL: Component 3 — Acknowledgments heading found but little text ({partial} pts)")
                total_score += partial
        else:
            print("FAIL: Component 3 — 'Acknowledgments' heading not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------
    # Component 4: Table of Contents with 5 chapters (0.20 points)
    # Must have "Table of Contents" heading and 5 chapter entries.
    # -----------------------------------------------------------------
    try:
        toc_idx = None
        for i, t in enumerate(all_para_texts_lower):
            if 'table of contents' in t:
                toc_idx = i
                break

        if toc_idx is not None:
            # Count chapter entries after TOC heading
            chapter_count = 0
            for j in range(toc_idx + 1, min(toc_idx + 15, len(all_para_texts))):
                t = all_para_texts_lower[j]
                if re.search(r'chapter\s+\d', t):
                    chapter_count += 1
                # Stop if we hit another section heading
                elif t in ('list of tables', 'list of figures', 'abstract',
                           'acknowledgments', 'acknowledgements'):
                    break

            if chapter_count >= 5:
                print(f"PASS: Component 4 — TOC with {chapter_count} chapters ({0.20} pts)")
                total_score += 0.20
            elif chapter_count >= 3:
                partial = 0.12
                print(f"PARTIAL: Component 4 — TOC found but only {chapter_count} chapters ({partial} pts)")
                total_score += partial
            elif chapter_count >= 1:
                partial = 0.06
                print(f"PARTIAL: Component 4 — TOC found but only {chapter_count} chapters ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — TOC heading found but no chapter entries")
        else:
            print("FAIL: Component 4 — 'Table of Contents' heading not found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -----------------------------------------------------------------
    # Component 5: List of Tables and List of Figures (0.10 points)
    # Both headings must be present in the document.
    # -----------------------------------------------------------------
    try:
        has_lot = any('list of tables' in t for t in all_para_texts_lower)
        has_lof = any('list of figures' in t for t in all_para_texts_lower)

        comp5 = 0.0
        if has_lot:
            comp5 += 0.05
            print("PASS: 'List of Tables' heading found")
        else:
            print("FAIL: 'List of Tables' heading not found")
        if has_lof:
            comp5 += 0.05
            print("PASS: 'List of Figures' heading found")
        else:
            print("FAIL: 'List of Figures' heading not found")

        if comp5 > 0:
            print(f"PASS: Component 5 — Lists section ({comp5:.2f} pts)")
            total_score += comp5
        else:
            print("FAIL: Component 5 — Neither List of Tables nor List of Figures found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -----------------------------------------------------------------
    # Component 6: Roman numeral page numbering (0.10 points)
    # At least one section must use roman numeral page number format.
    # Check for pgNumType with fmt=lowerRoman or upperRoman, or
    # instrText containing ROMANLOW/ROMANUPP.
    # -----------------------------------------------------------------
    try:
        roman_found = False
        for sec in doc.sections:
            sect_pr = sec._sectPr
            pgNumType = sect_pr.find(qn('w:pgNumType'))
            if pgNumType is not None:
                fmt = pgNumType.attrib.get(qn('w:fmt'), '')
                if 'roman' in fmt.lower():
                    roman_found = True
                    break

            # Also check footer instrText for roman format
            footer = sec.footer
            if footer and footer.paragraphs:
                for fp in footer.paragraphs:
                    for run in fp.runs:
                        instr_elems = run.element.findall(qn('w:instrText'))
                        for ie in instr_elems:
                            if ie.text and 'roman' in ie.text.lower():
                                roman_found = True
                                break
                    if roman_found:
                        break
            if roman_found:
                break

        if roman_found:
            print(f"PASS: Component 6 — Roman numeral page numbering found ({0.10} pts)")
            total_score += 0.10
        else:
            print("FAIL: Component 6 — No roman numeral page numbering found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
