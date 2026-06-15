"""
Initial Setup: Sociology research paper with 5 headings, no comments
Task ID: writer_struct_052
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_052'
OUTPUT = f'{WORKDIR}/Desktop/sociology_paper.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Set up document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ===== Title =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('Social Capital and Community Resilience: A Sociological Analysis')
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_run = author_para.add_run('Dr. Emily Hartwell, Department of Sociology, Westfield University')
    author_run.font.size = Pt(12)

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.add_run('March 2025')

    doc.add_page_break()

    # ===== Abstract =====
    abstract_heading = doc.add_paragraph('Abstract', style='Heading 2')
    abstract_text = doc.add_paragraph(
        'This paper examines the relationship between social capital formation and community resilience '
        'in post-industrial urban neighborhoods across three metropolitan areas in the United States. '
        'Drawing on a mixed-methods approach combining longitudinal survey data (n=1,247) with '
        'ethnographic fieldwork conducted between 2021 and 2024, the study identifies key mechanisms '
        'through which bonding and bridging social capital contribute to—or undermine—community adaptive '
        'capacity during periods of economic and environmental stress. Findings suggest that communities '
        'with higher levels of institutional trust and cross-class social ties demonstrate significantly '
        'greater resilience outcomes, even when controlling for material resource availability. '
        'Policy implications for urban planning and community development initiatives are discussed.'
    )

    doc.add_page_break()

    # ===== INTRODUCTION =====
    doc.add_heading('Introduction', level=1)

    doc.add_paragraph(
        'Community resilience has emerged as a central concern in contemporary sociology as urban '
        'populations face mounting pressures from deindustrialization, climate change, and widening '
        'economic inequality. While early resilience frameworks drew primarily from ecological systems '
        'theory (Holling, 1973; Walker et al., 2004), sociologists have increasingly argued that '
        'social structures, cultural resources, and institutional relationships play equally '
        'determinative roles in shaping how communities respond to and recover from adverse shocks.'
    )

    doc.add_paragraph(
        'Putnam\'s foundational work on social capital (1993, 2000) established the empirical link '
        'between associational life, civic engagement, and positive collective outcomes. Subsequent '
        'researchers have refined this framework, distinguishing between bonding capital—dense ties '
        'within homogeneous groups—and bridging capital—connections across social boundaries—and '
        'demonstrating their differential effects on community functioning (Briggs, 1998; '
        'Woolcock & Narayan, 2000). However, the specific pathways through which social capital '
        'translates into resilience capacity remain undertheorized and empirically contested.'
    )

    doc.add_paragraph(
        'This paper addresses this gap by examining social capital dynamics in three neighborhoods '
        'that have undergone significant economic disruption over the past decade: Millbrook Heights '
        'in Detroit, MI; Riverside Flats in Cleveland, OH; and Eastgate Commons in Baltimore, MD. '
        'Each site experienced substantial manufacturing job losses between 2010 and 2018, yet '
        'exhibited markedly different trajectories of recovery and community cohesion. By comparing '
        'these cases, we aim to identify the social structural conditions that enable or constrain '
        'resilient community responses to economic adversity.'
    )

    doc.add_paragraph(
        'The paper proceeds as follows: Section 2 reviews relevant theoretical background on social '
        'capital and community resilience. Section 3 describes our research methodology and data '
        'collection procedures. Section 4 presents empirical findings from both quantitative analyses '
        'and qualitative case studies. Section 5 discusses the implications of these findings for '
        'sociological theory and community development policy, and Section 6 concludes with '
        'recommendations for future research.'
    )

    doc.add_page_break()

    # ===== BACKGROUND =====
    doc.add_heading('Background', level=1)

    doc.add_paragraph(
        'The concept of social capital has a complex genealogy within sociological theory. '
        'Bourdieu (1986) first theorized social capital as resources accrued through membership '
        'in durable networks of mutual acquaintance and recognition, situating it within his broader '
        'framework of field theory and capital conversion. For Bourdieu, social capital was primarily '
        'a tool of class reproduction, enabling dominant groups to mobilize collective resources in '
        'ways that reinforced existing hierarchies.'
    )

    doc.add_paragraph(
        'Coleman (1988) offered a more functionalist account, defining social capital as those aspects '
        'of social structure that facilitate the actions of actors within that structure. His analysis '
        'emphasized closure—the density of ties within a network—as a key mechanism generating trust '
        'and normative compliance. Coleman\'s framework proved influential in educational sociology, '
        'where it explained differential academic outcomes between Catholic and public school students '
        'as a function of community-level social capital.'
    )

    doc.add_paragraph(
        'Putnam\'s (1993) landmark study of civic traditions in Italy operationalized social capital '
        'at the regional level, measuring it through indicators such as newspaper readership, '
        'electoral participation, and associational membership. His subsequent work on American '
        'civic decline (2000) traced a decades-long erosion of social capital across multiple '
        'domains, attributing this decline to generational change, suburbanization, and the rise '
        'of electronic entertainment.'
    )

    doc.add_paragraph(
        'Critics have challenged various aspects of this framework. Portes (1998) identified '
        'the "dark side" of social capital, noting that strong in-group solidarity can generate '
        'exclusionary pressures and restrict individual mobility. Bourgois (1995) documented how '
        'dense social networks in drug-economy neighborhoods could simultaneously provide material '
        'support and perpetuate cycles of violence and incarceration. These critiques underscore '
        'the importance of examining social capital in relation to broader structural inequalities '
        'rather than treating it as uniformly beneficial.'
    )

    doc.add_paragraph(
        'More recent scholarship has sought to integrate social capital frameworks with resilience '
        'theory. Norris et al. (2008) developed a comprehensive model of community resilience '
        'encompassing four primary adaptive capacities: economic development, social capital, '
        'information and communication, and community competence. Within this framework, social '
        'capital contributes to resilience through multiple pathways: enabling collective action, '
        'facilitating information sharing, providing emotional support, and fostering a sense of '
        'place attachment and collective identity.'
    )

    doc.add_page_break()

    # ===== METHODOLOGY =====
    doc.add_heading('Methodology', level=1)

    doc.add_paragraph(
        'This study employs a mixed-methods research design combining longitudinal survey research '
        'with multi-site ethnographic fieldwork. The integration of these approaches allows us to '
        'examine both the aggregate patterns of social capital distribution and the lived experiences '
        'and interpretive frameworks through which community members navigate periods of collective stress.'
    )

    doc.add_paragraph(
        'Survey data were collected in three waves (2021, 2022, 2024) using a stratified random '
        'sample drawn from residential address registries in each study neighborhood. The baseline '
        'sample comprised 1,247 adult residents (Millbrook Heights: n=423; Riverside Flats: n=398; '
        'Eastgate Commons: n=426). Attrition across waves was managed through active tracking '
        'procedures, resulting in a final longitudinal sample of 891 respondents (71.5% retention rate).'
    )

    doc.add_paragraph(
        'The survey instrument measured social capital through a multi-dimensional battery including: '
        'network size and composition (bonding vs. bridging ties); institutional trust in local '
        'government, police, schools, and religious organizations; civic participation in formal '
        'and informal associations; reciprocity norms; and collective efficacy as measured through '
        'Sampson et al.\'s (1997) validated scale. Resilience outcomes were assessed through '
        'indicators including employment stability, housing tenure, health status, and subjective '
        'wellbeing.'
    )

    doc.add_paragraph(
        'Ethnographic fieldwork was conducted by a team of four researchers, each embedded in '
        'one of the study neighborhoods for periods ranging from eight to fourteen months between '
        '2021 and 2024. Field methods included participant observation at community meetings, '
        'neighborhood associations, religious services, local businesses, and informal gathering '
        'spaces; semi-structured interviews with 148 community members, organizational leaders, '
        'and local officials; and document analysis of community planning materials, local news '
        'archives, and organizational records.'
    )

    doc.add_paragraph(
        'Quantitative data were analyzed using structural equation modeling to test mediation and '
        'moderation pathways between social capital dimensions and resilience outcomes. Qualitative '
        'data were analyzed through iterative thematic coding procedures drawing on grounded theory '
        'methodology (Charmaz, 2006). Mixed-methods integration occurred at the interpretation '
        'phase, with qualitative findings used to illuminate and contextualize patterns identified '
        'in quantitative analyses.'
    )

    doc.add_page_break()

    # ===== FINDINGS =====
    doc.add_heading('Findings', level=1)

    doc.add_paragraph(
        'Quantitative analyses revealed substantial variation in social capital levels across '
        'the three study neighborhoods and significant associations between specific social capital '
        'dimensions and resilience outcomes. Table 1 presents descriptive statistics for key '
        'variables at baseline. Eastgate Commons exhibited the highest mean levels of bridging '
        'social capital (M=3.74, SD=0.82), while Millbrook Heights showed the strongest bonding '
        'capital scores (M=4.12, SD=0.67). Riverside Flats fell between the other two sites '
        'on both dimensions.'
    )

    doc.add_paragraph(
        'Structural equation modeling identified institutional trust as the strongest direct '
        'predictor of resilience outcomes (β=0.43, p<0.001), followed by bridging social capital '
        '(β=0.31, p<0.001). Bonding social capital showed a more complex pattern: direct effects '
        'were non-significant, but significant positive indirect effects were observed through '
        'collective efficacy (β=0.18, p<0.01). Notably, the interaction between bonding capital '
        'and economic stress was negative (β=-0.22, p<0.01), suggesting that under conditions of '
        'severe resource scarcity, high levels of within-group solidarity may create insular dynamics '
        'that hinder adaptive response.'
    )

    doc.add_paragraph(
        'Longitudinal analyses revealed divergent trajectories across sites. In Eastgate Commons, '
        'bridging capital levels remained stable between 2021 and 2024, while resilience indicators '
        'showed modest improvement across the study period. Millbrook Heights exhibited declining '
        'bonding capital alongside deteriorating resilience outcomes, a pattern our ethnographic data '
        'suggest reflects the accelerating departure of long-term residents and organizational '
        'infrastructure. Riverside Flats showed the most complex pattern, with an initial decline '
        'in social capital followed by partial recovery associated with the emergence of new '
        'community organizations in 2022-2023.'
    )

    doc.add_paragraph(
        'Qualitative findings elaborated the mechanisms underlying these quantitative patterns. '
        'In Eastgate Commons, a dense network of cross-racial and cross-class civic associations—'
        'including a longstanding neighborhood development corporation, an interfaith council, '
        'and an active residents\' association—provided institutional infrastructure for collective '
        'problem-solving. Interviews with organizational leaders revealed deliberate strategies for '
        'cultivating bridging ties, including community events designed to attract residents across '
        'ethnic and class boundaries and mentorship programs connecting younger and older residents.'
    )

    doc.add_paragraph(
        'In Millbrook Heights, by contrast, the dominant organizations were highly homogeneous '
        'block clubs whose dense internal solidarity coexisted with weak ties to broader municipal '
        'and regional networks. When the neighborhood\'s largest employer closed its facility '
        'in 2019, these organizations lacked the external connections necessary to access '
        'information about alternative employment opportunities or to effectively advocate for '
        'community interests in city planning processes. The result was a collective sense of '
        'abandonment that further eroded trust in both local institutions and civic participation.'
    )

    doc.add_page_break()

    # ===== CONCLUSION =====
    doc.add_heading('Conclusion', level=1)

    doc.add_paragraph(
        'This study advances our understanding of the social structural conditions that enable '
        'community resilience in the face of economic adversity. Our findings suggest that the '
        'composition and configuration of social capital—particularly the balance between bonding '
        'and bridging ties and the level of institutional trust—matters as much as its overall '
        'volume in determining community adaptive capacity.'
    )

    doc.add_paragraph(
        'Theoretical contributions include the identification of institutional trust as a critical '
        'mediating variable linking social capital to resilience outcomes, and the specification of '
        'conditions under which bonding social capital may impede rather than facilitate community '
        'adaptation. These findings suggest the need for more nuanced models of social capital that '
        'account for structural position within broader institutional and political-economic contexts.'
    )

    doc.add_paragraph(
        'From a policy perspective, these results support investments in organizations and programs '
        'that deliberately cultivate bridging ties across social boundaries, rather than focusing '
        'exclusively on strengthening cohesion within existing community groups. Effective '
        'community development initiatives must also address the institutional trust deficits that '
        'constrain collective action in disadvantaged neighborhoods, through sustained commitments '
        'to responsive and accountable local governance.'
    )

    doc.add_paragraph(
        'Limitations of this study include its focus on a limited number of case sites and the '
        'challenges of establishing causal direction in the relationship between social capital '
        'and resilience outcomes. Future research should employ experimental or quasi-experimental '
        'designs where possible to strengthen causal inference, and should examine social capital '
        'dynamics in a wider range of geographic and demographic contexts. Additional work is '
        'also needed to explore how digital communication platforms and social media interact '
        'with traditional forms of community social capital to shape resilience trajectories '
        'in contemporary urban environments.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
