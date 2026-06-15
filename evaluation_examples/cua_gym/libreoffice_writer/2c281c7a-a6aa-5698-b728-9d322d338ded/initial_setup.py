"""
Initial Setup: 10-page audit report with empty footer
Task ID: writer_struct_069
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'audit_report'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_field(paragraph, field_instr: str):
    """Insert a Word field (e.g. DATE, PAGE, NUMPAGES) into a paragraph run."""
    from docx.oxml import OxmlElement

    # fldChar begin
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run_begin._r.append(fld_begin)

    # instrText
    run_instr = paragraph.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = field_instr
    run_instr._r.append(instr)

    # fldChar end
    run_end = paragraph.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run_end._r.append(fld_end)


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Header (simple title) ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "Meridian Financial Group — Internal Audit Division"
    hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in hp.runs:
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x47, 0x2C)

    # --- Footer: empty paragraph (no fields, no text) ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()  # ensure footer paragraph is empty

    # ============================================================
    # Document body — 10-page realistic audit report
    # ============================================================

    # Title
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(24)
    title_para.paragraph_format.space_after = Pt(12)
    run_t = title_para.add_run("INTERNAL AUDIT REPORT")
    run_t.bold = True
    run_t.font.size = Pt(18)

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_para.paragraph_format.space_after = Pt(6)
    run_s = subtitle_para.add_run("Fiscal Year 2024 — Compliance & Financial Controls Review")
    run_s.italic = True
    run_s.font.size = Pt(13)

    meta_para = doc.add_paragraph()
    meta_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta_para.paragraph_format.space_after = Pt(24)
    meta_run = meta_para.add_run("Prepared by: Internal Audit Department\nDate of Issue: January 15, 2025\nReport Reference: IAD-2024-Q4-001")
    meta_run.font.size = Pt(11)

    doc.add_page_break()

    # ----- Page 2: Executive Summary -----
    h1 = doc.add_heading("1. Executive Summary", level=1)
    h1.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        "This report presents the findings of the annual internal audit conducted by the Internal Audit "
        "Department (IAD) of Meridian Financial Group for fiscal year 2024. The audit covered four primary "
        "domains: financial reporting integrity, procurement and vendor management, information technology "
        "general controls, and regulatory compliance. A total of 147 audit procedures were executed across "
        "12 business units between October 1, 2024 and December 20, 2024."
    )
    doc.add_paragraph(
        "Overall, the organization demonstrated satisfactory compliance with established policies and "
        "applicable regulations. However, three significant findings and nine moderate findings were identified "
        "requiring management attention. Recommendations have been provided for each finding to strengthen the "
        "control environment and mitigate associated risks."
    )

    h2 = doc.add_heading("1.1 Scope and Objectives", level=2)
    doc.add_paragraph(
        "The primary objectives of this audit were: (1) to assess the adequacy and effectiveness of internal "
        "controls over financial reporting; (2) to evaluate compliance with corporate policies, regulatory "
        "requirements, and contractual obligations; (3) to identify operational inefficiencies and recommend "
        "process improvements; and (4) to verify the safeguarding of company assets."
    )

    doc.add_page_break()

    # ----- Page 3: Methodology -----
    h1 = doc.add_heading("2. Audit Methodology", level=1)
    h1.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        "The audit was conducted in accordance with the International Standards for the Professional Practice "
        "of Internal Auditing (IPPF) issued by the Institute of Internal Auditors (IIA). The audit team "
        "comprised four senior auditors, two IT audit specialists, and one data analytics expert."
    )

    h2 = doc.add_heading("2.1 Risk Assessment", level=2)
    doc.add_paragraph(
        "A risk-based approach was adopted to prioritize audit areas. Prior to fieldwork, the IAD conducted "
        "a comprehensive risk assessment involving interviews with 23 senior managers and analysis of prior "
        "audit findings, regulatory examination reports, and key performance indicators. Risk ratings (High, "
        "Medium, Low) were assigned to each audit universe item and used to determine the scope and depth "
        "of testing procedures."
    )

    h2 = doc.add_heading("2.2 Data Analytics", level=2)
    doc.add_paragraph(
        "For the first time, the IAD deployed continuous auditing tools to analyze 100% of procurement "
        "transactions (approximately 48,200 transactions totaling $312.7 million) and payroll disbursements "
        "(24,816 payments totaling $89.4 million). Automated exception reports flagged anomalies for "
        "targeted review, significantly improving audit coverage and efficiency."
    )

    doc.add_page_break()

    # ----- Page 4: Financial Reporting -----
    h1 = doc.add_heading("3. Financial Reporting & Controls", level=1)
    h1.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        "The audit assessed key controls over the financial close process, journal entry preparation, "
        "reconciliations, and management review procedures. Testing covered the period January 1 through "
        "September 30, 2024."
    )

    h2 = doc.add_heading("3.1 Journal Entry Controls", level=2)
    doc.add_paragraph(
        "A sample of 340 manual journal entries was selected for testing from a population of 14,782 "
        "entries. Testing focused on authorization, completeness of supporting documentation, and "
        "appropriate account coding. Results indicate that 98.2% of sampled entries complied with policy. "
        "Six entries lacked proper authorization documentation; management has committed to reinforcing "
        "approval workflows by Q1 2025."
    )

    # Add a simple table
    doc.add_paragraph()
    tbl_heading = doc.add_paragraph("Table 1: Journal Entry Testing Summary")
    tbl_heading.runs[0].bold = True
    tbl_heading.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    headers_row = table.rows[0]
    for idx, hdr_text in enumerate(["Control Area", "Sample Size", "Exception Rate"]):
        cell = headers_row.cells[idx]
        cell.text = hdr_text
        cell.paragraphs[0].runs[0].bold = True

    data_rows = [
        ("Authorization", "340", "1.8%"),
        ("Supporting Docs", "340", "0.9%"),
        ("Account Coding", "340", "0.3%"),
        ("Period-End Review", "120", "0.0%"),
    ]
    for r_idx, (c1, c2, c3) in enumerate(data_rows, 1):
        table.rows[r_idx].cells[0].text = c1
        table.rows[r_idx].cells[1].text = c2
        table.rows[r_idx].cells[2].text = c3

    doc.add_page_break()

    # ----- Page 5: Procurement -----
    h1 = doc.add_heading("4. Procurement & Vendor Management", level=1)
    h1.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        "Procurement activities were reviewed across all business units. Key areas examined included vendor "
        "selection and onboarding, purchase order compliance, invoice matching, and vendor performance monitoring."
    )

    h2 = doc.add_heading("4.1 Vendor Onboarding", level=2)
    doc.add_paragraph(
        "Of 87 new vendors engaged in 2024, a sample of 25 were reviewed for completeness of onboarding "
        "documentation including W-9/W-8 forms, certificate of insurance, conflict-of-interest declarations, "
        "and credit references. Three vendors (12%) lacked current insurance certificates. The Procurement "
        "team has initiated an annual recertification process to address this gap."
    )

    h2 = doc.add_heading("4.2 Purchase Order Compliance", level=2)
    doc.add_paragraph(
        "Analysis of 48,200 transactions identified 213 instances (0.44%) where purchases exceeded the "
        "delegated approval threshold without appropriate escalation. Transaction values ranged from $1,200 "
        "to $78,500. Affected business units included Operations (94 instances), IT Infrastructure (67 "
        "instances), and Facilities Management (52 instances). Management has acknowledged the finding and "
        "will implement system-enforced workflow controls by March 2025."
    )

    doc.add_page_break()

    # ----- Page 6: IT Controls -----
    h1 = doc.add_heading("5. Information Technology General Controls", level=1)
    h1.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        "IT General Controls (ITGCs) were assessed for five in-scope applications: the core banking system "
        "(CoreBanx v12.4), the ERP platform (SAP S/4HANA 2023), the payroll system (ADP Workforce Now), "
        "the document management system (SharePoint Online), and the treasury management system (Kyriba)."
    )

    h2 = doc.add_heading("5.1 Access Management", level=2)
    doc.add_paragraph(
        "User access reviews were conducted for all five systems. A total of 3,847 active user accounts "
        "were reviewed. Findings included: 142 accounts belonging to terminated employees that had not been "
        "deprovisioned within the 24-hour policy requirement (average delay: 4.2 business days); 28 accounts "
        "with excessive privileges inconsistent with job function; and 6 shared/generic accounts in "
        "production environments contrary to policy."
    )

    h2 = doc.add_heading("5.2 Change Management", level=2)
    doc.add_paragraph(
        "A sample of 60 system changes (40 normal, 20 emergency) was tested for adherence to the change "
        "management policy. All 40 normal changes complied with required approvals, testing documentation, "
        "and implementation verification. Of the 20 emergency changes, 3 (15%) lacked post-implementation "
        "documentation as required by policy. The IT department has committed to updating the emergency "
        "change process by end of Q1 2025."
    )

    doc.add_page_break()

    # ----- Page 7: HR & Payroll -----
    h1 = doc.add_heading("6. Human Resources & Payroll", level=1)
    h1.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        "The payroll audit covered 24,816 payment transactions across 1,247 active employees. "
        "Total payroll expenditure for the audit period was $89.4 million. Testing procedures included "
        "verification of hiring authorization, salary change approvals, overtime calculations, and "
        "benefit deduction accuracy."
    )

    h2 = doc.add_heading("6.1 Payroll Reconciliation", level=2)
    doc.add_paragraph(
        "Monthly payroll reconciliations between the payroll system and general ledger were reviewed for "
        "ten months. All reconciliations were completed within required timeframes. Minor variances totaling "
        "$4,287 were identified and properly investigated and cleared. No significant unexplained variances "
        "were noted."
    )

    h2 = doc.add_heading("6.2 Overtime and Supplemental Pay", level=2)
    doc.add_paragraph(
        "A targeted review of overtime payments revealed that 14 employees in the Warehouse Operations "
        "division received overtime pay exceeding 30% of base salary in at least two months during the "
        "period. While not in violation of policy, this pattern warrants management review to assess whether "
        "staffing levels are adequate to meet operational requirements without excessive overtime reliance."
    )

    doc.add_page_break()

    # ----- Page 8: Regulatory Compliance -----
    h1 = doc.add_heading("7. Regulatory Compliance", level=1)
    h1.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        "The compliance review assessed adherence to applicable laws, regulations, and internal policies "
        "including the Bank Secrecy Act (BSA), Anti-Money Laundering (AML) requirements, the Sarbanes-Oxley "
        "Act (SOX) Section 302 and 404 obligations, GDPR data privacy requirements, and the company's own "
        "Code of Business Conduct and Ethics."
    )

    h2 = doc.add_heading("7.1 BSA/AML Compliance", level=2)
    doc.add_paragraph(
        "The BSA/AML compliance program was reviewed with particular focus on Currency Transaction Reports "
        "(CTR), Suspicious Activity Reports (SAR), and Customer Due Diligence (CDD) documentation. "
        "Testing of 50 CDD files revealed that 4 files (8%) for high-risk customers lacked required "
        "enhanced due diligence documentation updates. These files have been escalated to the BSA Officer "
        "for immediate remediation."
    )

    h2 = doc.add_heading("7.2 SOX Compliance", level=2)
    doc.add_paragraph(
        "Key controls supporting the SOX 302 and 404 certifications were tested. All 47 key controls were "
        "operational and effective during the testing period. The external auditors, Deloitte & Touche LLP, "
        "have concurred with management's assessment and are expected to issue an unqualified opinion on "
        "internal controls over financial reporting."
    )

    doc.add_page_break()

    # ----- Page 9: Summary of Findings -----
    h1 = doc.add_heading("8. Summary of Findings and Recommendations", level=1)
    h1.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        "This section consolidates all audit findings identified during the engagement. Findings are "
        "categorized by severity: Significant (potential material impact), Moderate (notable control "
        "weakness), and Advisory (best practice recommendation)."
    )

    # Findings table
    tbl_heading2 = doc.add_paragraph("Table 2: Audit Findings Summary")
    tbl_heading2.runs[0].bold = True
    tbl_heading2.paragraph_format.space_after = Pt(4)

    findings_table = doc.add_table(rows=13, cols=4)
    findings_table.style = "Table Grid"
    fh = findings_table.rows[0]
    for idx, fhdr in enumerate(["Ref.", "Finding Description", "Severity", "Owner"]):
        fh.cells[idx].text = fhdr
        fh.cells[idx].paragraphs[0].runs[0].bold = True

    findings_data = [
        ("F-01", "Terminated employee accounts not deprovisioned timely", "Significant", "IT Security"),
        ("F-02", "Purchase order approval threshold violations", "Significant", "Procurement"),
        ("F-03", "CDD documentation gaps for high-risk customers", "Significant", "BSA Officer"),
        ("F-04", "Vendor insurance certificates not current", "Moderate", "Procurement"),
        ("F-05", "Journal entries lacking authorization documentation", "Moderate", "Controller"),
        ("F-06", "Excessive user privileges in production systems", "Moderate", "IT Security"),
        ("F-07", "Emergency change management documentation gaps", "Moderate", "IT Change Mgmt"),
        ("F-08", "Excessive overtime in Warehouse Operations", "Advisory", "HR / Operations"),
        ("F-09", "Shared accounts in production environments", "Moderate", "IT Security"),
        ("F-10", "Vendor performance KPIs not formally tracked", "Advisory", "Procurement"),
        ("F-11", "Outdated BCM/DR plan documentation", "Moderate", "IT Operations"),
        ("F-12", "Missing annual policy attestations (17 employees)", "Advisory", "Compliance"),
    ]
    for r_idx, row_vals in enumerate(findings_data, 1):
        for c_idx, val in enumerate(row_vals):
            findings_table.rows[r_idx].cells[c_idx].text = val

    doc.add_page_break()

    # ----- Page 10: Management Responses & Conclusion -----
    h1 = doc.add_heading("9. Management Responses", level=1)
    h1.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        "Management has reviewed all findings and provided the following responses. All significant and "
        "moderate findings have been accepted with remediation timelines. Advisory findings will be addressed "
        "as resources permit."
    )

    doc.add_paragraph(
        "F-01 (IT Security): The Identity & Access Management team will implement automated deprovisioning "
        "integration with the HR system by February 28, 2025. Interim manual review process established "
        "effective January 20, 2025. Responsible: Jennifer Park, CISO."
    )
    doc.add_paragraph(
        "F-02 (Procurement): System-enforced purchase order workflow controls will be deployed in SAP by "
        "March 31, 2025. Retroactive review of all flagged transactions completed with no financial loss "
        "identified. Responsible: David Okonkwo, VP Procurement."
    )
    doc.add_paragraph(
        "F-03 (BSA Officer): All 4 identified CDD files have been escalated; enhanced due diligence "
        "completed for 3 files as of report date. Fourth file under review. Responsible: Linda Vasquez, BSA Officer."
    )

    h1 = doc.add_heading("10. Conclusion", level=1)
    h1.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        "Based on the results of our audit procedures, it is our opinion that the overall system of internal "
        "controls at Meridian Financial Group is adequate and generally effective. The findings identified "
        "in this report do not individually or collectively represent a material weakness in the organization's "
        "control environment; however, timely remediation is strongly recommended to maintain and strengthen "
        "the control posture."
    )
    doc.add_paragraph(
        "The Internal Audit Department will conduct follow-up procedures in Q2 2025 to verify that all "
        "significant and moderate findings have been remediated in accordance with management's committed "
        "action plans. This report has been provided to the Audit Committee of the Board of Directors for "
        "their review and consideration."
    )

    sign_para = doc.add_paragraph()
    sign_para.paragraph_format.space_before = Pt(24)
    sign_para.add_run("Robert T. Harrington, CIA, CRMA\n").bold = True
    sign_para.add_run("Chief Audit Executive\n")
    sign_para.add_run("Meridian Financial Group\n")
    sign_para.add_run("January 15, 2025")

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
