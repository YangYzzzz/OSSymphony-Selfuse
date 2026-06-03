"""
Initial Setup: Financial analysis report with embedded revenue data task
Task ID: osworld_multi_apps_calc_to_writer_004
Domain: libreoffice_writer (multi-app: also creates an xlsx for Calc)

Creates:
  1. ~/Documents/finance/annual_revenue.xlsx  — quarterly revenue data (Q1-Q4 2023, multiple regions)
  2. /home/user/osworld_multi_apps_calc_to_writer_004.docx — financial analysis doc with 'Revenue Analysis'
     section but NO table yet (agent must add Q4 2023 rows as table)
Opens LibreOffice Writer (docx) and LibreOffice Calc (xlsx) for the agent.
"""

import os
import shlex
import subprocess
import time

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_calc_to_writer_004'
DOCX_OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
XLSX_DIR = f'{WORKDIR}/Documents/finance'
XLSX_OUTPUT = f'{XLSX_DIR}/annual_revenue.xlsx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch a GUI app on the VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_xlsx():
    """Create ~/Documents/finance/annual_revenue.xlsx with quarterly revenue data."""
    os.makedirs(XLSX_DIR, exist_ok=True)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Revenue Data"

    # Header row
    headers = ["Quarter", "Region", "Revenue", "Growth"]
    header_font = Font(bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="FF2E4057", end_color="FF2E4057", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center")

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    # Revenue data — realistic multi-region quarterly data for 2023
    data = [
        # Q1 2023
        ("Q1 2023", "North America",  4823500,  0.082),
        ("Q1 2023", "Europe",         3142800,  0.065),
        ("Q1 2023", "Asia Pacific",   2587300,  0.114),
        ("Q1 2023", "Latin America",   892100,  0.043),
        ("Q1 2023", "Middle East",     614700,  0.097),
        # Q2 2023
        ("Q2 2023", "North America",  5104200,  0.058),
        ("Q2 2023", "Europe",         3389600,  0.079),
        ("Q2 2023", "Asia Pacific",   2961400,  0.144),
        ("Q2 2023", "Latin America",   978500,  0.097),
        ("Q2 2023", "Middle East",     703200,  0.144),
        # Q3 2023
        ("Q3 2023", "North America",  5432100,  0.064),
        ("Q3 2023", "Europe",         3621000,  0.069),
        ("Q3 2023", "Asia Pacific",   3198700,  0.080),
        ("Q3 2023", "Latin America",  1056400,  0.080),
        ("Q3 2023", "Middle East",     788900,  0.122),
        # Q4 2023
        ("Q4 2023", "North America",  6187400,  0.139),
        ("Q4 2023", "Europe",         4052300,  0.119),
        ("Q4 2023", "Asia Pacific",   3745800,  0.171),
        ("Q4 2023", "Latin America",  1234600,  0.169),
        ("Q4 2023", "Middle East",     956200,  0.212),
    ]

    thin = Side(style="thin", color="CCCCCC")
    data_border = Border(left=thin, right=thin, top=thin, bottom=thin)

    alt_fill_a = PatternFill(start_color="FFF7F9FC", end_color="FFF7F9FC", fill_type="solid")
    alt_fill_b = PatternFill(start_color="FFEEF2F8", end_color="FFEEF2F8", fill_type="solid")

    for row_idx, (quarter, region, revenue, growth) in enumerate(data, 2):
        fill = alt_fill_a if row_idx % 2 == 0 else alt_fill_b
        vals = [quarter, region, revenue, growth]
        for col_idx, val in enumerate(vals, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.border = data_border
            cell.fill = fill
            if col_idx == 3:
                cell.number_format = '$#,##0'
                cell.alignment = Alignment(horizontal="right")
            elif col_idx == 4:
                cell.number_format = '0.0%'
                cell.alignment = Alignment(horizontal="right")
            else:
                cell.alignment = Alignment(horizontal="left")

    # Column widths
    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 10

    # Freeze header row
    ws.freeze_panes = "A2"

    wb.save(XLSX_OUTPUT)
    print(f"Spreadsheet created: {XLSX_OUTPUT}")


def create_docx():
    """Create financial_analysis.docx with a 'Revenue Analysis' section but no table yet."""
    doc = Document()

    # Title
    title = doc.add_heading("Annual Financial Analysis Report — FY 2023", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle / metadata
    subtitle = doc.add_paragraph("Prepared by: Finance Strategy & Intelligence Team")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_para = doc.add_paragraph("Report Date: February 14, 2024")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.runs[0].font.italic = True
    date_para.runs[0].font.size = Pt(11)
    date_para.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ---- Executive Summary ----
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Fiscal year 2023 marked a period of sustained growth for the organization, with total consolidated "
        "revenue reaching $47.8 million — an increase of 11.3% over FY 2022. Growth was broad-based across "
        "all regions, with particularly strong performance in the Asia Pacific corridor and a record-breaking "
        "fourth quarter driven by new enterprise contract closures and favorable foreign-exchange tailwinds."
    )
    doc.add_paragraph(
        "Operating margins improved by 140 basis points year-over-year, reaching 22.6%. The Finance team "
        "attributes this efficiency gain to the successful integration of the Madrid and Singapore operations "
        "completed in Q2, which eliminated approximately $1.2 million in duplicated overhead costs."
    )

    # ---- Business Environment ----
    doc.add_heading("Business Environment", level=1)
    doc.add_paragraph(
        "The macroeconomic backdrop for FY 2023 was characterized by moderating inflation, central-bank "
        "rate stabilization in North America and Europe, and continued post-pandemic demand recovery in "
        "emerging markets. These tailwinds contributed to higher enterprise software spending and accelerated "
        "digital-transformation project timelines for many of our key accounts."
    )
    doc.add_paragraph(
        "Key competitive dynamics: three of the organization's top-ten strategic accounts expanded their "
        "contract scope in H2 2023. Two greenfield wins in the Middle East corridor contributed $0.9 million "
        "in incremental revenue for the year and position the region for further growth in 2024."
    )

    # ---- Revenue Analysis section (NO table yet — agent must add it) ----
    doc.add_heading("Revenue Analysis", level=1)
    doc.add_paragraph(
        "This section provides a detailed breakdown of revenue by quarter and geographic region for fiscal "
        "year 2023. Data has been sourced from the consolidated finance data warehouse and reconciled against "
        "audited general-ledger balances as of December 31, 2023."
    )
    doc.add_paragraph(
        "Full year revenue grew 11.3% on a reported basis and 9.8% on a constant-currency basis. Q4 2023 "
        "represented the strongest quarter, reflecting typical seasonality as well as an accelerated deal-close "
        "cycle driven by year-end budget utilization by enterprise customers. The Q4 2023 revenue detail by "
        "region is available in ~/Documents/finance/annual_revenue.xlsx and should be embedded below for "
        "reference."
    )
    # NOTE: No table here — agent needs to add one

    # ---- Cost Structure ----
    doc.add_heading("Cost Structure and Margins", level=1)
    doc.add_paragraph(
        "Total operating expenses for FY 2023 were $37.0 million, compared to $33.9 million in FY 2022. "
        "The increase was primarily driven by headcount growth in product engineering (+28 FTEs) and "
        "expansion of the go-to-market team in Asia Pacific (+14 FTEs). Sales & Marketing spend as a "
        "percentage of revenue declined from 31% to 29%, reflecting improved sales productivity."
    )

    cost_breakdown = [
        ("Cost Category", "FY 2022 ($M)", "FY 2023 ($M)", "YoY Change"),
        ("Cost of Revenue", "12.4", "13.8", "+11.3%"),
        ("Research & Development", "8.1", "9.6", "+18.5%"),
        ("Sales & Marketing", "10.8", "11.9", "+10.2%"),
        ("General & Administrative", "2.6", "1.7", "-34.6%"),
        ("Total OpEx", "33.9", "37.0", "+9.1%"),
    ]
    cost_table = doc.add_table(rows=len(cost_breakdown), cols=4)
    cost_table.style = "Table Grid"
    for r_idx, row_data in enumerate(cost_breakdown):
        for c_idx, val in enumerate(row_data):
            cell = cost_table.cell(r_idx, c_idx)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            if r_idx == 0:
                run.bold = True

    # ---- Outlook ----
    doc.add_heading("Outlook and Forward Guidance", level=1)
    doc.add_paragraph(
        "Management provides initial FY 2024 guidance of revenue in the range of $52.5–$54.0 million, "
        "representing growth of 10–13% over FY 2023. This guidance assumes continued strength in North "
        "America and Asia Pacific, moderate recovery in European enterprise spending following a soft H1 "
        "2023, and the full-year contribution of two recently announced partnership agreements."
    )
    doc.add_paragraph(
        "Operating margin is expected to expand modestly to 23.5–24.5% as fixed-cost leverage improves "
        "and the prior-year investments in headcount begin to generate productivity returns. Capital "
        "expenditure is budgeted at approximately $2.1 million, focused primarily on data-center "
        "infrastructure upgrades and security compliance initiatives."
    )

    doc.save(DOCX_OUTPUT)
    print(f"Writer document created: {DOCX_OUTPUT}")


def main():
    create_xlsx()
    create_docx()

    # GUI-ready startup — open the Writer document (primary task surface)
    launch_gui(f'libreoffice --writer "{DOCX_OUTPUT}"', delay_sec=3.0)
    # Open the spreadsheet so the agent can inspect/copy Q4 2023 data
    launch_gui(f'libreoffice --calc "{XLSX_OUTPUT}"', delay_sec=2.0)

    print("GUI_READY: launched LibreOffice Writer and LibreOffice Calc with DISPLAY=:0")


main()
