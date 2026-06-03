"""
Reward Script: Insert Q4 2023 revenue data as a table in the Revenue Analysis section
Task ID: osworld_multi_apps_calc_to_writer_004
Domain: libreoffice_writer (multi-app: reads from calc, inserts into writer)

Scoring Rubric:
  Component 1 (0.4 pts): A new table with Q4 2023 revenue data exists in the document
                          (document has a table that is NOT the Cost Category table)
  Component 2 (0.3 pts): The new revenue table has the correct 4-column headers
                          (Quarter, Region, Revenue, Growth)
  Component 3 (0.3 pts): The new revenue table is positioned in the Revenue Analysis
                          section (after "Revenue Analysis" heading, before
                          "Cost Structure and Margins" heading in body order)

Why these components:
  - Initial env has 1 table (Cost Structure/Cost Category table only)
  - Golden env has 2 tables: Q4 2023 revenue table + Cost Category table
  - The task specifically asks to embed Q4 2023 data in the Revenue Analysis section
  - All three components FAIL on initial and PASS on golden
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_calc_to_writer_004'

# Expected Q4 2023 data from annual_revenue.xlsx
EXPECTED_HEADERS = ['Quarter', 'Region', 'Revenue', 'Growth']
EXPECTED_Q4_REGIONS = ['North America', 'Europe', 'Asia Pacific', 'Latin America', 'Middle East']

# Known Q4 2023 revenue values from annual_revenue.xlsx
EXPECTED_Q4_DATA = {
    'North America': (6187400, 0.139),
    'Europe': (4052300, 0.119),
    'Asia Pacific': (3745800, 0.171),
    'Latin America': (1234600, 0.169),
    'Middle East': (956200, 0.212),
}


def persist_app_state():
    """Attempt to save any unsaved Writer edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def find_revenue_table(doc):
    """
    Find a table in the document that appears to contain Q4 2023 revenue data.
    Returns the table object or None.
    A revenue table is identified by having 'Quarter' or 'Q4 2023' in first column
    and NOT having 'Cost Category' in first column.
    """
    for table in doc.tables:
        if not table.rows:
            continue
        first_row_texts = [cell.text.strip() for cell in table.rows[0].cells]
        # This is the revenue/quarterly table (not the cost category table)
        if len(first_row_texts) >= 1:
            first_cell = first_row_texts[0].strip()
            if first_cell in ('Quarter', 'Q4 2023'):
                return table
    return None


def get_body_element_order(doc):
    """
    Return an ordered list of (type, text_snippet) for body elements.
    type is 'para' or 'table'.
    """
    order = []
    body = doc.element.body
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for child in body:
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'p':
            text = ''.join(n.text or '' for n in child.iter() if n.tag.endswith('}t'))
            style_elem = child.find(f'{{{ns}}}pPr/{{{ns}}}pStyle')
            style = style_elem.get(f'{{{ns}}}val') if style_elem is not None else 'Normal'
            order.append(('para', style, text))
        elif tag == 'tbl':
            rows = child.findall(f'.//{{{ns}}}tr')
            cols_first_row = rows[0].findall(f'.//{{{ns}}}tc') if rows else []
            first_text = ''.join(
                n.text or '' for n in cols_first_row[0].iter() if n.tag.endswith('}t')
            ) if cols_first_row else ''
            order.append(('table', '', first_text))
    return order


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — gate check
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # --- Component 1: A Q4 2023 revenue table exists in the document (0.4 points) ---
    # The initial doc has only 1 table (Cost Category). The golden doc has 2 tables.
    # We check that a table with 'Quarter' or 'Q4 2023' header exists and contains
    # Q4 2023 data rows.
    try:
        revenue_table = find_revenue_table(doc)
        if revenue_table is not None:
            # Verify at least some Q4 2023 rows are present
            q4_rows_found = 0
            for row in revenue_table.rows[1:]:  # skip header
                cells = [cell.text.strip() for cell in row.cells]
                if cells and 'Q4 2023' in cells[0]:
                    q4_rows_found += 1

            if q4_rows_found > 0:
                print(f"PASS: Component 1 — Q4 2023 revenue table found with {q4_rows_found} Q4 data rows (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 1 — Revenue table found but contains no Q4 2023 rows")
        else:
            print("FAIL: Component 1 — No Q4 2023 revenue table found in document")
            # Check how many tables exist for diagnostic info
            print(f"  Diagnostic: document has {len(doc.tables)} table(s)")
            for i, t in enumerate(doc.tables):
                if t.rows:
                    print(f"  Table {i} first row: {[c.text.strip() for c in t.rows[0].cells]}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # --- Component 2: Revenue table has correct 4-column headers (0.3 points) ---
    # Headers must be: Quarter, Region, Revenue, Growth
    try:
        revenue_table = find_revenue_table(doc)
        if revenue_table is not None and revenue_table.rows:
            header_cells = [cell.text.strip() for cell in revenue_table.rows[0].cells]
            # Check all 4 expected headers are present
            headers_match = all(h in header_cells for h in EXPECTED_HEADERS)
            if headers_match and len(header_cells) == 4:
                print(f"PASS: Component 2 — Revenue table has correct headers {header_cells} (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Expected headers {EXPECTED_HEADERS}, found: {header_cells}")
        else:
            print("FAIL: Component 2 — No revenue table available to check headers")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # --- Component 3: Revenue table is placed in the Revenue Analysis section (0.3 points) ---
    # The table must appear AFTER the 'Revenue Analysis' heading and BEFORE
    # the 'Cost Structure and Margins' heading in body element order.
    try:
        order = get_body_element_order(doc)

        revenue_section_idx = None
        cost_section_idx = None
        revenue_table_idx = None

        for i, item in enumerate(order):
            elem_type, style, text = item
            if elem_type == 'para' and style in ('Heading1', 'Heading 1'):
                if 'Revenue Analysis' in text and revenue_section_idx is None:
                    revenue_section_idx = i
                elif 'Cost Structure' in text and cost_section_idx is None:
                    cost_section_idx = i
            elif elem_type == 'table' and text.strip() in ('Quarter', 'Q4 2023'):
                revenue_table_idx = i

        # Evaluate placement: table must be between Revenue Analysis and Cost Structure headings
        all_found = (revenue_section_idx is not None and
                     cost_section_idx is not None and
                     revenue_table_idx is not None)
        if not all_found:
            missing = []
            if revenue_section_idx is None:
                missing.append("'Revenue Analysis' heading")
            if cost_section_idx is None:
                missing.append("'Cost Structure' heading")
            if revenue_table_idx is None:
                missing.append("revenue table in body")
            print(f"FAIL: Component 3 — Missing elements: {', '.join(missing)}")
        elif revenue_section_idx < revenue_table_idx < cost_section_idx:
            print(f"PASS: Component 3 — Revenue table correctly placed in Revenue Analysis section "
                  f"(body positions: RevenueHeading={revenue_section_idx}, Table={revenue_table_idx}, "
                  f"CostHeading={cost_section_idx}) (0.3 pts)")
            if revenue_section_idx < revenue_table_idx < cost_section_idx:
                total_score += 0.3
        else:
            print(f"FAIL: Component 3 — Revenue table NOT in Revenue Analysis section. "
                  f"Body positions: RevenueHeading={revenue_section_idx}, "
                  f"Table={revenue_table_idx}, CostHeading={cost_section_idx}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point — persist first, then verify
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
