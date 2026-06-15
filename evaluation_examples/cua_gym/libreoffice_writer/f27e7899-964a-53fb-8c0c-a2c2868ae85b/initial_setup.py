"""
Initial Setup: Set up document with gutter margin for binding
Task ID: writer_rd_054
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_054'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set standard margins: 2.54 cm on all sides (no gutter)
    for section in doc.sections:
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.page_width = Cm(21.0)   # A4
        section.page_height = Cm(29.7)  # A4

    section = doc.sections[0]

    # --- Title Page ---
    title = doc.add_heading('Network Infrastructure Operations Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies Inc.')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    version = doc.add_paragraph()
    version.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = version.add_run('Version 3.2 — March 2025')
    run.font.size = Pt(12)
    run.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    confidential = doc.add_paragraph()
    confidential.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = confidential.add_run('CONFIDENTIAL — Internal Use Only')
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_page_break()

    # --- Table of Contents Page ---
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1. Introduction and Scope', '3'),
        ('2. Network Architecture Overview', '5'),
        ('3. Hardware Inventory and Specifications', '8'),
        ('4. IP Address Management', '12'),
        ('5. VLAN Configuration Standards', '15'),
        ('6. Firewall Rules and Security Policies', '18'),
        ('7. VPN and Remote Access Procedures', '22'),
        ('8. Monitoring and Alerting Framework', '25'),
        ('9. Backup and Disaster Recovery', '29'),
        ('10. Incident Response Procedures', '33'),
        ('11. Change Management Protocol', '37'),
        ('12. Vendor Contact Information', '40'),
        ('13. Compliance and Audit Requirements', '43'),
        ('14. Appendices', '47'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(f'{item} ').font.size = Pt(11)
        dots = '.' * (60 - len(item))
        p.add_run(f'{dots} {page}').font.size = Pt(11)

    doc.add_page_break()

    # --- Chapter content generation ---
    chapters = [
        {
            'title': 'Introduction and Scope',
            'sections': [
                ('Purpose', [
                    'This manual provides comprehensive documentation for all network infrastructure operations at Meridian Technologies Inc. It covers the design, configuration, maintenance, and troubleshooting procedures for the corporate network spanning three data centers and fourteen branch offices across North America.',
                    'The primary audience includes network engineers, system administrators, and IT operations staff responsible for maintaining uptime and performance targets as defined in the corporate SLA framework.',
                    'All procedures documented herein have been reviewed and approved by the Network Architecture Review Board (NARB) as of the latest revision date. Any deviations from documented procedures require written approval from the Director of Network Operations, currently Dr. Patricia Alvarez.',
                ]),
                ('Scope and Applicability', [
                    'This document applies to all production, staging, and development network environments managed by the Infrastructure Operations team. Test lab environments maintained by individual engineering teams are excluded unless they connect to production network segments.',
                    'Branch office networks covered include: San Francisco (HQ), Seattle, Portland, Denver, Dallas, Chicago, Atlanta, Miami, New York, Boston, Toronto, Vancouver, Phoenix, and Austin. Each branch maintains a standardized network stack as described in Section 3.',
                ]),
                ('Document Conventions', [
                    'Throughout this manual, commands intended for execution on Cisco IOS devices are shown in monospace font. Configuration file excerpts are enclosed in bordered boxes. Warning notes are prefixed with a caution symbol and highlighted in amber.',
                    'IP addresses shown in examples use the documentation ranges (192.0.2.0/24, 198.51.100.0/24, 203.0.113.0/24) per RFC 5737 unless referring to actual production allocations, which are listed in the IP Address Management section.',
                ]),
            ]
        },
        {
            'title': 'Network Architecture Overview',
            'sections': [
                ('Core Network Topology', [
                    'The Meridian Technologies network follows a three-tier hierarchical design consisting of core, distribution, and access layers. The core layer comprises redundant Cisco Nexus 9500 switches deployed in a spine-leaf fabric across our primary data center in San Francisco and the disaster recovery site in Dallas.',
                    'Distribution layer switches aggregate traffic from access layer switches and apply policy-based routing, QoS marking, and inter-VLAN routing. Each distribution pair operates in an active-active VRRP configuration to eliminate single points of failure.',
                    'The access layer provides endpoint connectivity through Cisco Catalyst 9300 series switches supporting 802.1X port-based authentication, dynamic VLAN assignment via RADIUS, and Power over Ethernet Plus (PoE+) for IP phones and wireless access points.',
                ]),
                ('WAN Architecture', [
                    'Branch office connectivity is provided through a hybrid WAN architecture combining MPLS circuits from AT&T Business with SD-WAN overlay tunnels using Cisco Viptela controllers. Each branch maintains a primary 500 Mbps MPLS circuit and a secondary 1 Gbps broadband internet connection for SD-WAN failover.',
                    'Traffic engineering policies prioritize voice and video conferencing traffic over the MPLS backbone, while bulk data transfers and software updates are directed over the internet-based SD-WAN tunnels. The SD-WAN controller cluster at HQ manages centralized policy deployment across all branch sites.',
                ]),
                ('Cloud Connectivity', [
                    'Direct connections to AWS (us-west-2, us-east-1) and Azure (West US 2, East US) are maintained through AWS Direct Connect and Azure ExpressRoute circuits terminating at the San Francisco and Dallas data centers respectively. Each cloud connection provides 10 Gbps dedicated bandwidth with BGP peering.',
                    'Cloud security groups and network ACLs are managed through Terraform Infrastructure-as-Code templates stored in the internal GitLab repository. Changes to cloud network configurations follow the same change management process described in Section 11.',
                ]),
            ]
        },
        {
            'title': 'Hardware Inventory and Specifications',
            'sections': [
                ('Core Switches', [
                    'The core network fabric consists of four Cisco Nexus 9504 chassis switches, two at each data center location. Each chassis is populated with N9K-X97160YC-EX line cards providing 48 ports of 10/25 Gbps and 4 ports of 40/100 Gbps. The supervisor modules run NX-OS version 10.3(2) with EVPN-VXLAN fabric configuration.',
                    'Hardware specifications per chassis: Switching capacity 17.6 Tbps, forwarding rate 8.28 Bpps, redundant power supplies (3000W AC), redundant fabric modules. Mean time between failure (MTBF) rated at 320,000 hours per Cisco reliability data.',
                ]),
                ('Distribution Switches', [
                    'Distribution layer switches are Cisco Catalyst 9500-48Y4C models deployed in stacked pairs at each location. Stack interconnects use 100 Gbps QSFP28 cables. Each stack supports up to 1,200 MAC addresses per port and 32,000 total MAC entries in hardware.',
                    'Current firmware: IOS-XE 17.9.4a. Firmware updates are scheduled quarterly in coordination with Cisco TAC recommendations and internal vulnerability assessment findings. Emergency patches may be applied outside the regular schedule following the expedited change process.',
                ]),
                ('Wireless Infrastructure', [
                    'Wireless coverage is provided by Cisco Catalyst 9130AXI access points managed by redundant Cisco 9800-CL wireless controllers running in SSO (Stateful Switchover) mode. The HQ campus has 847 access points across 12 floors, providing seamless roaming via 802.11r fast transition.',
                    'RF design targets a minimum signal strength of -67 dBm at 5 GHz band across all work areas. Channel planning uses Cisco AI-enhanced RRM (Radio Resource Management) with DCA (Dynamic Channel Assignment) enabled. The 6 GHz band (Wi-Fi 6E) is being rolled out in phases starting with executive conference rooms and engineering labs.',
                ]),
            ]
        },
        {
            'title': 'IP Address Management',
            'sections': [
                ('Address Allocation Strategy', [
                    'Meridian Technologies uses a hierarchical IP addressing scheme based on geographic region and function. The corporate allocation is 10.0.0.0/8 subdivided as follows: 10.1.0.0/16 for San Francisco HQ, 10.2.0.0/16 for Dallas DR site, 10.10.0.0/16 through 10.23.0.0/16 for branch offices in sequential order.',
                    'Within each /16 allocation, subnets are further divided by function: /24 blocks for user VLANs, /27 blocks for management networks, /28 blocks for point-to-point links, and /30 blocks for loopback addresses. DHCP pools are configured on distribution layer switches with lease times of 8 hours for wired clients and 4 hours for wireless clients.',
                ]),
                ('DNS Architecture', [
                    'Internal DNS is hosted on redundant Windows Server 2022 domain controllers with Active Directory-integrated zones. Forward lookup zones cover meridiantech.local and all subdomain delegations. Reverse lookup zones are maintained for all /16 allocations.',
                    'External DNS is hosted on AWS Route 53 with health check-based failover for public-facing services. DNSSEC is enabled on all external zones. DNS query logging is forwarded to the Splunk SIEM for security monitoring and threat intelligence correlation.',
                ]),
                ('IPv6 Transition Plan', [
                    'IPv6 deployment follows a dual-stack approach starting with the data center core and extending to distribution and access layers over a 24-month timeline. The allocated IPv6 prefix is 2001:db8:cafe::/48, with /64 subnets assigned per VLAN. Link-local addressing uses EUI-64 for infrastructure devices.',
                    'Phase 1 (completed): Core router and firewall IPv6 configuration. Phase 2 (in progress): Distribution layer dual-stack with OSPFv3 and BGP IPv6 address families. Phase 3 (planned Q3 2025): Access layer and endpoint IPv6 provisioning via SLAAC with RDNSS.',
                ]),
            ]
        },
        {
            'title': 'VLAN Configuration Standards',
            'sections': [
                ('VLAN Naming Convention', [
                    'All VLANs follow the naming convention: SITE-FUNCTION-ID. For example, SF-DATA-100 represents the San Francisco data VLAN with ID 100. Voice VLANs use the range 200-299, management VLANs use 900-999, and guest networks are assigned VLAN 666 at all sites for consistency.',
                    'Native VLAN is configured as VLAN 999 (MGMT-NATIVE) across all trunk links. The default VLAN 1 is administratively disabled on all switch ports. Unused ports are placed in VLAN 998 (QUARANTINE) and shut down administratively.',
                ]),
                ('Trunk Configuration', [
                    'Inter-switch trunk links use 802.1Q encapsulation with explicit VLAN allowed lists. No trunk port permits all VLANs; each trunk is configured with only the VLANs required for that specific link. Dynamic Trunking Protocol (DTP) is disabled globally on all switches.',
                    'Trunk pruning is enabled to optimize broadcast domain containment. VTP mode is set to transparent on all switches to prevent unintended VLAN propagation. Each site maintains its own VLAN database managed through Ansible automation playbooks.',
                ]),
                ('Access Port Standards', [
                    'Access ports are configured with the following security features: BPDU Guard enabled, Root Guard on uplinks, PortFast on all edge ports, DHCP Snooping with rate limiting set to 15 packets per second, and Dynamic ARP Inspection on all user VLANs.',
                    'Storm control thresholds are set at 80% for broadcast, 80% for multicast, and 80% for unicast with the action set to shutdown. Err-disabled recovery is configured with a 300-second timer for BPDU Guard and DHCP rate limit violations. All violations are logged to the central syslog server at 10.1.0.50.',
                ]),
            ]
        },
        {
            'title': 'Firewall Rules and Security Policies',
            'sections': [
                ('Perimeter Firewall Architecture', [
                    'The perimeter security stack consists of Palo Alto PA-5260 firewalls deployed in active-passive HA pairs at each data center. The firewalls enforce zone-based security policies across five defined zones: EXTERNAL, DMZ, INTERNAL, MANAGEMENT, and CLOUD-TRANSIT.',
                    'All inter-zone traffic requires explicit allow rules; the default policy is deny-all with logging. Rule base optimization is performed monthly using the Palo Alto Expedition tool to identify unused and shadowed rules. The current rule base contains approximately 2,400 active rules across both data centers.',
                ]),
                ('Intrusion Prevention', [
                    'Threat prevention profiles are applied to all inter-zone security policies. The IPS signature database is updated every 4 hours from the Palo Alto threat intelligence feed. Custom IPS signatures are maintained for internal applications identified during annual penetration testing.',
                    'SSL/TLS decryption is enabled for outbound internet traffic from user VLANs with exceptions for financial services, healthcare portals, and employee benefits sites as required by the privacy policy. Decryption certificates are managed by the PKI team and renewed annually.',
                ]),
                ('Microsegmentation', [
                    'Data center workloads are protected by VMware NSX-T distributed firewall rules implementing zero-trust microsegmentation. Security groups are defined by application tier (web, app, database) and environment (production, staging, development).',
                    'East-west traffic policies are managed through NSX-T policy objects synchronized with the CMDB. New application deployments must include a microsegmentation policy template as part of the architecture review process documented in Section 11.',
                ]),
            ]
        },
        {
            'title': 'VPN and Remote Access Procedures',
            'sections': [
                ('Site-to-Site VPN', [
                    'Site-to-site VPN tunnels between branch offices and data centers use IKEv2 with certificate-based authentication. Certificates are issued by the internal Microsoft ADCS Certificate Authority with a validity period of 2 years. Phase 1 parameters: AES-256-GCM encryption, SHA-384 integrity, DH Group 20 (384-bit ECP).',
                    'Phase 2 (IPsec) parameters: ESP with AES-256-GCM, Perfect Forward Secrecy using DH Group 20, and a rekeying interval of 28,800 seconds. Dead peer detection is configured with a 10-second interval and 5 retries before declaring the tunnel down and initiating failover.',
                ]),
                ('Remote Access VPN', [
                    'Remote workers connect via Cisco AnyConnect Secure Mobility Client version 5.0 or later. The VPN concentrator is a Cisco ASA 5585-X deployed in the DMZ zone. Authentication uses SAML 2.0 integration with Okta for single sign-on, followed by Cisco Duo push notification for multi-factor authentication.',
                    'Split tunneling is enabled by default, routing only corporate network prefixes (10.0.0.0/8, 172.16.0.0/12) through the VPN tunnel. Internet-bound traffic exits directly from the remote endpoint. Full tunnel mode is available upon request for users handling sensitive data classifications.',
                ]),
                ('VPN Monitoring', [
                    'VPN tunnel status is monitored continuously by the SolarWinds Network Performance Monitor (NPM) with 60-second polling intervals. Tunnel latency, jitter, and packet loss are tracked against SLA thresholds: latency below 80ms, jitter below 30ms, and packet loss below 0.1%.',
                    'Automated alerts are configured for tunnel flaps (more than 3 state changes in 10 minutes), latency threshold breaches, and certificate expiration warnings at 30, 14, and 7 days before expiry. The Network Operations Center (NOC) maintains a 24/7 staffed response capability.',
                ]),
            ]
        },
        {
            'title': 'Monitoring and Alerting Framework',
            'sections': [
                ('Monitoring Stack Overview', [
                    'The monitoring infrastructure is built on a three-tier architecture: data collection (Telegraf agents, SNMP polling, NetFlow/sFlow), data storage and processing (InfluxDB time-series database, Elasticsearch cluster), and visualization and alerting (Grafana dashboards, PagerDuty integration).',
                    'SNMP v3 polling is configured with authentication (SHA-256) and privacy (AES-256) on all network devices. Polling intervals are 60 seconds for interface utilization, 300 seconds for CPU/memory metrics, and 30 seconds for critical path availability checks.',
                ]),
                ('Alert Escalation Matrix', [
                    'Alerts are classified into four severity levels. P1 (Critical): Complete service outage affecting multiple users, response time 15 minutes, escalation to VP of Engineering if unresolved in 1 hour. P2 (High): Degraded service or single point of failure loss, response time 30 minutes.',
                    'P3 (Medium): Non-critical component failure with redundancy still intact, response time 4 hours during business hours. P4 (Low): Informational alerts and capacity planning triggers, response time next business day. Severity classification is determined automatically by PagerDuty based on alert source and affected service catalog entries.',
                ]),
                ('Dashboard Standards', [
                    'Standard Grafana dashboards are maintained for each network tier and service. The NOC overview dashboard displays aggregate health scores for all sites using a red/amber/green traffic light system. Drill-down dashboards provide per-site and per-device views with 7-day rolling trend analysis.',
                    'Dashboard templates are version controlled in the GitLab infrastructure repository. Changes to production dashboards follow the standard change request process. Each dashboard includes a documentation panel describing the data sources, thresholds, and intended audience.',
                ]),
            ]
        },
        {
            'title': 'Backup and Disaster Recovery',
            'sections': [
                ('Network Device Backup', [
                    'Running configurations for all network devices are backed up automatically every 6 hours using Oxidized (open-source network configuration backup tool). Backups are stored in a dedicated GitLab repository with full version history and diff comparison capability.',
                    'Configuration compliance checks run daily comparing active configurations against the approved baseline templates. Deviations are reported to the Network Engineering team via email digest and Slack channel notification. Non-compliant devices are flagged for remediation within 48 hours.',
                ]),
                ('Disaster Recovery Procedures', [
                    'The network DR plan defines two scenarios: partial site failure (single component or service) and complete site failure (entire data center or branch office). Recovery Time Objective (RTO) for the core network is 4 hours; Recovery Point Objective (RPO) for configuration data is 6 hours.',
                    'Complete site failover from San Francisco to Dallas is tested semi-annually during planned maintenance windows. The failover process involves BGP route manipulation to shift traffic to the DR site, DNS TTL reduction 24 hours before the test, and activation of the DR site firewall policy set.',
                ]),
                ('Backup Verification', [
                    'Monthly backup restoration tests verify that the most recent configurations can be successfully loaded onto spare hardware. The test environment includes one device from each platform type (Nexus, Catalyst, Palo Alto, ASA) maintained in the lab rack at the San Francisco data center.',
                    'Test results are documented in the Configuration Management Database (CMDB) with pass/fail status and any noted discrepancies. Failed restoration tests trigger an immediate review of the backup process and corrective action within 72 hours.',
                ]),
            ]
        },
        {
            'title': 'Incident Response Procedures',
            'sections': [
                ('Incident Classification', [
                    'Network incidents are classified using the ITIL v4 incident management framework adapted for Meridian Technologies. Incidents are categorized by affected service, impacted user count, and business criticality score. The ServiceNow ITSM platform is the system of record for all incident tickets.',
                    'Major incidents (P1/P2) trigger the Major Incident Management (MIM) process, which includes immediate conference bridge activation, stakeholder notification within 15 minutes, and executive status updates every 30 minutes until resolution. Post-incident reviews are mandatory within 5 business days.',
                ]),
                ('Troubleshooting Procedures', [
                    'Network troubleshooting follows a structured approach: verify physical connectivity, check interface status and error counters, validate Layer 2 forwarding (MAC table, ARP table), verify Layer 3 routing (routing table, BGP neighbors), and test application-layer connectivity.',
                    'Common diagnostic tools include: ping/traceroute for path validation, show interface/show ip route for device state, packet capture (SPAN/ERSPAN) for traffic analysis, and NetFlow data for historical traffic pattern review. All troubleshooting steps must be documented in the incident ticket.',
                ]),
                ('Communication Templates', [
                    'Standard communication templates are maintained in the ServiceNow knowledge base for initial notification, status updates, and resolution notifications. Templates include placeholders for: affected service name, impact description, estimated time to resolution, workaround instructions, and root cause summary.',
                    'External customer notifications for service-affecting incidents are coordinated through the Customer Success team. Internal notifications are distributed via the IT Service Status page, email distribution lists, and Slack channels (#it-status, #network-alerts).',
                ]),
            ]
        },
        {
            'title': 'Change Management Protocol',
            'sections': [
                ('Change Request Process', [
                    'All network changes must be submitted through the ServiceNow Change Management module at least 5 business days before the planned implementation date. Emergency changes follow an expedited process with verbal CAB approval followed by retrospective documentation within 24 hours.',
                    'Change requests must include: detailed description of the change, risk assessment score (using the standard 5x5 risk matrix), rollback plan with specific commands, estimated implementation duration, affected services and user groups, and required maintenance window if applicable.',
                ]),
                ('Change Advisory Board', [
                    'The Change Advisory Board meets weekly on Tuesdays at 2:00 PM Pacific to review pending change requests. CAB members include representatives from Network Engineering, Security Operations, Application Support, and the Service Desk. Quorum requires at least one representative from each team.',
                    'Standard changes (pre-approved, low-risk changes following documented procedures) are auto-approved and do not require CAB review. The current standard change catalog includes: VLAN creation/modification, firewall rule additions for approved applications, and switch port configuration changes.',
                ]),
                ('Post-Implementation Review', [
                    'All changes undergo a post-implementation review within 2 business days of completion. The review verifies that the change achieved its intended objective, no unintended side effects occurred, monitoring confirms normal operation, and documentation has been updated.',
                    'Failed changes or changes causing unplanned service impact are escalated to a formal Problem Management investigation. Root cause analysis follows the 5-Whys methodology with findings documented in the Problem record and shared during the monthly IT Operations review meeting.',
                ]),
            ]
        },
        {
            'title': 'Vendor Contact Information',
            'sections': [
                ('Primary Vendors', [
                    'Cisco Systems — TAC Support: 1-800-553-2447, Contract ID: MERID-ENT-2024-0847, Escalation Manager: David Park (dpark@cisco.com). Smart Net Total Care coverage on all production equipment. Response time: 2-hour for P1, 4-hour for P2, NBD for P3/P4.',
                    'Palo Alto Networks — Support Portal: support.paloaltonetworks.com, Customer ID: C-2840173, Premium Support with Platinum SLA. Dedicated Technical Account Manager: Lisa Fernandez (lfernandez@paloaltonetworks.com), monthly review calls scheduled first Wednesday.',
                ]),
                ('Circuit Providers', [
                    'AT&T Business — MPLS circuits, Account Team Lead: Robert Kim (robert.kim@att.com), Account Number: 847-2940-3857. Circuit IDs documented in the CMDB per-site records. Maintenance notifications sent to noc-maintenance@meridiantech.com.',
                    'Zayo Group — Dark fiber and wavelength services between data centers, Account Manager: Maria Santos (msantos@zayo.com). 100 Gbps wavelength between SF and Dallas with path diversity. Contract renewal date: December 2026.',
                ]),
                ('Software and Services', [
                    'SolarWinds — Network Performance Monitor and Network Configuration Manager licenses. Support portal: customerportal.solarwinds.com, License key: SW-MERID-2024-NPM-500. Annual maintenance renewal in March.',
                    'Splunk Enterprise — SIEM and log management platform. Cloud deployment on Splunk Cloud. Account Executive: James Chen (jchen@splunk.com). Daily ingest volume: approximately 850 GB. Contract includes 24/7 Premium support.',
                ]),
            ]
        },
        {
            'title': 'Compliance and Audit Requirements',
            'sections': [
                ('Regulatory Framework', [
                    'Meridian Technologies network infrastructure must comply with SOC 2 Type II, PCI DSS v4.0 (for payment processing network segments), and HIPAA Security Rule (for healthcare client data environments). Annual third-party audits are conducted by Deloitte with interim self-assessments performed quarterly.',
                    'Network segmentation between compliance zones is enforced at the firewall level with dedicated security zones for PCI and HIPAA environments. Inter-zone traffic is logged and retained for 12 months in compliance with audit requirements. Access to compliance zones requires additional authorization through the GRC (Governance, Risk, and Compliance) team.',
                ]),
                ('Audit Logging', [
                    'All network device authentication events, configuration changes, and administrative commands are logged to the centralized Splunk SIEM. Log retention policy: 90 days hot storage (searchable), 1 year warm storage (archived), 7 years cold storage (compliance archive) for PCI-scoped devices.',
                    'Syslog transport uses TLS-encrypted connections to prevent log tampering in transit. Log integrity is verified using SHA-256 hash chains computed daily. Any gaps in log collection trigger automated alerts to the Security Operations Center.',
                ]),
                ('Audit Preparation Checklist', [
                    'Pre-audit activities begin 30 days before the scheduled audit date. The Network Engineering team must: verify all device firmware is within the approved version matrix, confirm NTP synchronization across all devices (stratum 3 or better), validate that unused ports are disabled and placed in the quarantine VLAN.',
                    'Documentation deliverables include: current network diagrams (physical and logical), firewall rule base with business justification for each rule, access control lists with last review dates, vulnerability scan reports from the most recent quarterly scan, and evidence of change management compliance for all changes in the audit period.',
                ]),
            ]
        },
        {
            'title': 'Appendices',
            'sections': [
                ('Appendix A: Standard Configuration Templates', [
                    'This appendix contains the approved baseline configuration templates for each network device platform. Templates are maintained in the GitLab repository at infrastructure/network/templates/ and are deployed via Ansible automation. All templates include mandatory security hardening settings.',
                    'Template versions are tagged using semantic versioning (MAJOR.MINOR.PATCH). The current approved versions are: Nexus 9500 template v2.4.1, Catalyst 9300 template v3.1.0, Catalyst 9500 template v2.2.3, PA-5260 template v1.8.0, and ASA 5585 template v2.0.1.',
                ]),
                ('Appendix B: Emergency Contact Procedures', [
                    'After-hours emergency support is provided by the on-call Network Engineer rotation. On-call schedule is published monthly in PagerDuty and the IT shared calendar. On-call engineers must be reachable within 15 minutes and capable of remote access to all network management systems.',
                    'Escalation path: On-call Engineer (15 min) → Network Engineering Manager (30 min) → Director of Network Operations (1 hour) → VP of Engineering (2 hours). For life-safety or physical security incidents, contact Facilities Management at ext. 5555 and Corporate Security at ext. 5500 simultaneously.',
                ]),
                ('Appendix C: Acronym Glossary', [
                    'ACL: Access Control List. ARP: Address Resolution Protocol. BGP: Border Gateway Protocol. BPDU: Bridge Protocol Data Unit. CAB: Change Advisory Board. CIDR: Classless Inter-Domain Routing. CMDB: Configuration Management Database. DCA: Dynamic Channel Assignment.',
                    'DHCP: Dynamic Host Configuration Protocol. DMZ: Demilitarized Zone. DNS: Domain Name System. DTP: Dynamic Trunking Protocol. EVPN: Ethernet VPN. GRE: Generic Routing Encapsulation. HA: High Availability. IDS: Intrusion Detection System. IKE: Internet Key Exchange.',
                    'IPS: Intrusion Prevention System. IPsec: Internet Protocol Security. MPLS: Multiprotocol Label Switching. NAT: Network Address Translation. NTP: Network Time Protocol. OSPF: Open Shortest Path First. QoS: Quality of Service. RADIUS: Remote Authentication Dial-In User Service.',
                    'RRM: Radio Resource Management. SLA: Service Level Agreement. SNMP: Simple Network Management Protocol. SPAN: Switched Port Analyzer. SSH: Secure Shell. TLS: Transport Layer Security. VLAN: Virtual Local Area Network. VPN: Virtual Private Network. VRRP: Virtual Router Redundancy Protocol. VXLAN: Virtual Extensible LAN.',
                ]),
            ]
        },
    ]

    for ch_idx, chapter in enumerate(chapters):
        doc.add_heading(f'{ch_idx + 1}. {chapter["title"]}', level=1)

        for sec_title, paragraphs in chapter['sections']:
            doc.add_heading(sec_title, level=2)
            for para_text in paragraphs:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(6)

        # Add a page break after each chapter except the last
        if ch_idx < len(chapters) - 1:
            doc.add_page_break()

    # Ensure all sections have the correct margins (in case new sections were created)
    for section in doc.sections:
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
