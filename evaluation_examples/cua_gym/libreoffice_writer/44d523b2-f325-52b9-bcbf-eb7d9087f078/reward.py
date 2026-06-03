"""
Reward Script: Insert cover page for appellate brief
Task ID: writer_legal_076
Domain: libreoffice_writer
Scoring:
  Component 1: Document has 2+ sections (section break separating cover page) — 0.15
  Component 2: Court name centered at top in bold — 0.25
  Component 3: Case caption centered in middle — 0.20
  Component 4: "APPELLANT'S OPENING BRIEF" centered, bold, 14pt — 0.20
  Component 5: Attorney info block at bottom, left-aligned — 0.10
  Component 6: Cover page section has no page numbering in footer — 0.10
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_076'


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice state."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify cover page insertion for appellate brief.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ========================================================
    # Component 1: Document has 2+ sections (0.15 points)
    # The cover page should be in its own section, separated by a section break.
    # Initial file has only 1 section.
    # ========================================================
    try:
        num_sections = len(doc.sections)
        if num_sections >= 2:
            print(f"PASS: Component 1 — Document has {num_sections} sections (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Expected 2+ sections, found {num_sections}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ========================================================
    # Component 2: Court name centered at top in bold (0.25 points)
    # Golden: First paragraphs contain "IN THE COURT OF APPEAL",
    # "OF THE STATE OF CALIFORNIA", "SECOND APPELLATE DISTRICT"
    # all centered and bold.
    # Initial: First paragraph is "TABLE OF CONTENTS" heading.
    # ========================================================
    try:
        court_keywords = ["COURT OF APPEAL", "STATE OF CALIFORNIA", "APPELLATE DISTRICT"]
        found_count = 0
        centered_bold_count = 0

        # Check first 5 paragraphs for court name lines
        for p in doc.paragraphs[:5]:
            text = p.text.strip().upper()
            for kw in court_keywords:
                if kw in text:
                    found_count += 1
                    is_centered = p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                    is_bold = any(r.bold for r in p.runs if r.text.strip())
                    if is_centered and is_bold:
                        centered_bold_count += 1
                    break

        if found_count >= 3 and centered_bold_count >= 3:
            print(f"PASS: Component 2 — Court name found ({found_count} lines), all centered+bold (0.25 pts)")
            total_score += 0.25
        elif found_count >= 2 and centered_bold_count >= 2:
            print(f"PARTIAL: Component 2 — Court name partially found ({found_count}/{centered_bold_count}) (0.15 pts)")
            total_score += 0.15
        elif found_count >= 1:
            print(f"PARTIAL: Component 2 — Court name partially found ({found_count} lines) (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 2 — No court name found in first 5 paragraphs")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ========================================================
    # Component 3: Case caption centered in middle (0.20 points)
    # Golden has "ELENA VASQUEZ", "Plaintiff and Appellant", "v.",
    # "PACIFIC COAST INDUSTRIES", centered.
    # Initial starts with TABLE OF CONTENTS — no such caption.
    # ========================================================
    try:
        caption_keywords = ["VASQUEZ", "APPELLANT", "PACIFIC COAST", "RESPONDENT"]
        found_caption = 0
        centered_caption = 0

        # Scan paragraphs 5-20 for caption elements
        for p in doc.paragraphs[3:20]:
            text = p.text.strip().upper()
            for kw in caption_keywords:
                if kw in text:
                    found_caption += 1
                    if p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                        centered_caption += 1
                    break

        if found_caption >= 3 and centered_caption >= 3:
            print(f"PASS: Component 3 — Case caption found ({found_caption} elements), {centered_caption} centered (0.20 pts)")
            total_score += 0.20
        elif found_caption >= 2:
            print(f"PARTIAL: Component 3 — Caption partially found ({found_caption} elements) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — No case caption found (found {found_caption} elements)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ========================================================
    # Component 4: "APPELLANT'S OPENING BRIEF" centered, bold, 14pt (0.20 points)
    # This is a key element of the cover page. Must be centered, bold, ~14pt.
    # Initial file has no such text.
    # ========================================================
    try:
        found_brief_title = False
        brief_centered = False
        brief_bold = False
        brief_size_ok = False

        for p in doc.paragraphs[:20]:
            text = p.text.strip().upper()
            if "APPELLANT" in text and "OPENING BRIEF" in text:
                found_brief_title = True
                brief_centered = p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                brief_bold = any(r.bold for r in p.runs if r.text.strip())
                for r in p.runs:
                    if r.text.strip() and r.font.size:
                        if abs(r.font.size.pt - 14.0) < 1.0:
                            brief_size_ok = True
                break

        if found_brief_title and brief_centered and brief_bold and brief_size_ok:
            print(f"PASS: Component 4 — 'APPELLANT'S OPENING BRIEF' centered, bold, 14pt (0.20 pts)")
            total_score += 0.20
        elif found_brief_title and brief_centered and brief_bold:
            print(f"PARTIAL: Component 4 — Title found centered+bold but size={brief_size_ok} (0.15 pts)")
            total_score += 0.15
        elif found_brief_title and brief_centered:
            print(f"PARTIAL: Component 4 — Title found centered but bold={brief_bold} (0.10 pts)")
            total_score += 0.10
        elif found_brief_title:
            print(f"PARTIAL: Component 4 — Title found but not properly formatted (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 4 — 'APPELLANT'S OPENING BRIEF' not found in first 20 paragraphs")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ========================================================
    # Component 5: Attorney info block at bottom left (0.10 points)
    # Golden has law firm name, attorney name, address, etc. left-aligned.
    # These do not exist in the initial file's first section.
    # ========================================================
    try:
        attorney_keywords = ["ATTORNEY", "COUNSEL", "LAW OFFICE", "SBN", "SUITE", "TELEPHONE"]
        found_attorney = 0
        left_aligned = 0

        # Look in paragraphs before the section break (first 30 paragraphs)
        for p in doc.paragraphs[:30]:
            text = p.text.strip().upper()
            for kw in attorney_keywords:
                if kw in text:
                    found_attorney += 1
                    if p.paragraph_format.alignment in (WD_PARAGRAPH_ALIGNMENT.LEFT, None):
                        left_aligned += 1
                    break

        if found_attorney >= 3 and left_aligned >= 2:
            print(f"PASS: Component 5 — Attorney info found ({found_attorney} lines), {left_aligned} left-aligned (0.10 pts)")
            total_score += 0.10
        elif found_attorney >= 1:
            print(f"PARTIAL: Component 5 — Some attorney info found ({found_attorney} lines) (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 — No attorney information found on cover page")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # ========================================================
    # Component 6: Cover page has no page number (0.10 points)
    # Golden: Section 0 footer has no page number field codes, or
    # the first page has suppressed header/footer.
    # Initial: Only 1 section with no specific first-page handling.
    # We check that the cover page section's footer has no PAGE field.
    # ========================================================
    try:
        if len(doc.sections) >= 2:
            first_section = doc.sections[0]
            # Check default footer of section 0 for PAGE field codes
            footer = first_section.footer
            footer_xml = footer._element.xml if footer._element is not None else ""

            has_page_field = "PAGE" in footer_xml or "w:fldChar" in footer_xml

            # Also check first page footer
            try:
                first_page_ftr = first_section.first_page_footer
                first_ftr_xml = first_page_ftr._element.xml if first_page_ftr._element is not None else ""
                has_page_field = has_page_field or ("PAGE" in first_ftr_xml and "w:fldChar" in first_ftr_xml)
            except:
                pass

            if not has_page_field:
                print(f"PASS: Component 6 — Cover page section has no page number in footer (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 6 — Cover page section footer contains page number field")
        else:
            print(f"FAIL: Component 6 — Cannot check (only 1 section, no separate cover page)")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
