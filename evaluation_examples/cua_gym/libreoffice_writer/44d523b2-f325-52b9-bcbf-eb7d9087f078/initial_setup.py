"""
Initial Setup: Appellate brief without cover page
Task ID: writer_legal_076
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_076'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard letter size with 1-inch margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Brief content starts directly on page 1 (NO cover page) ---

    # Table of Contents heading
    h1 = doc.add_heading('TABLE OF CONTENTS', level=1)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    toc_items = [
        ('Table of Authorities', 'ii'),
        ('Statement of the Case', '1'),
        ('Statement of Facts', '3'),
        ('Standard of Review', '7'),
        ('Argument', '9'),
        ('    I. The Trial Court Erred in Granting Summary Judgment', '9'),
        ('    II. The Evidence Was Insufficient to Support the Verdict', '15'),
        ('    III. The Court Abused Its Discretion in Excluding Key Testimony', '21'),
        ('Conclusion', '28'),
        ('Certificate of Compliance', '29'),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{item}')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run2 = p.add_run(f' {"." * (50 - len(item))} {page}')
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'

    doc.add_page_break()

    # Table of Authorities
    h2 = doc.add_heading('TABLE OF AUTHORITIES', level=1)
    h2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    cases = [
        'Aguilar v. Atlantic Richfield Co. (2001) 25 Cal.4th 826',
        'Buss v. Superior Court (1997) 16 Cal.4th 35',
        'Celotex Corp. v. Catrett (1986) 477 U.S. 317',
        'Doe v. City of Los Angeles (2007) 42 Cal.4th 531',
        'Evidence Code section 352',
        'Guz v. Bechtel National, Inc. (2000) 24 Cal.4th 317',
        'Kelly v. New West Federal Savings (1996) 49 Cal.App.4th 659',
        'People v. Watson (1956) 46 Cal.2d 818',
        'Sargon Enterprises, Inc. v. University of Southern California (2012) 55 Cal.4th 747',
    ]
    for case in cases:
        p = doc.add_paragraph()
        run = p.add_run(case)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.italic = True

    doc.add_page_break()

    # Statement of the Case
    h3 = doc.add_heading('STATEMENT OF THE CASE', level=1)
    h3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    case_paras = [
        'This appeal arises from a wrongful termination action filed by Appellant Elena Vasquez '
        'against Respondent Pacific Coast Industries, Inc. ("PCI"). Ms. Vasquez was employed by '
        'PCI as a Senior Quality Assurance Engineer from March 2019 until her termination on '
        'September 14, 2024.',

        'On November 2, 2024, Appellant filed a complaint in the Los Angeles County Superior '
        'Court, Case No. BC-2024-78432, alleging causes of action for (1) wrongful termination '
        'in violation of public policy; (2) retaliation under Labor Code section 1102.5; '
        '(3) breach of the implied covenant of good faith and fair dealing; and (4) intentional '
        'infliction of emotional distress.',

        'Respondent filed a motion for summary judgment on February 28, 2025. After briefing '
        'and oral argument, the Honorable Judge Patricia Nakamura granted the motion in its '
        'entirety on May 15, 2025. The court entered judgment in favor of Respondent on '
        'June 3, 2025.',

        'Appellant timely filed her Notice of Appeal on June 28, 2025.',
    ]

    for text in case_paras:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    doc.add_page_break()

    # Statement of Facts
    h4 = doc.add_heading('STATEMENT OF FACTS', level=1)
    h4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    facts_paras = [
        'Ms. Vasquez joined PCI in March 2019 as a Quality Assurance Engineer and was promoted '
        'to Senior Quality Assurance Engineer in January 2021. (Clerk\'s Transcript ["CT"] '
        'at 45.) Throughout her tenure, she consistently received exemplary performance reviews, '
        'including ratings of "Exceeds Expectations" for fiscal years 2020 through 2023. '
        '(CT at 48-52.)',

        'In April 2024, Ms. Vasquez discovered that PCI was systematically underreporting '
        'product defect rates to the California Department of Consumer Affairs. Specifically, '
        'she identified that the QA database had been modified to exclude approximately 340 '
        'defect reports filed between January and March 2024. (CT at 78-82; Reporter\'s '
        'Transcript ["RT"] at 234:15-236:8.)',

        'On April 22, 2024, Ms. Vasquez reported her findings to her direct supervisor, '
        'Robert Chen, Director of Quality Assurance. Mr. Chen instructed Ms. Vasquez to '
        '"forget about it" and stated that "corporate handles the reporting." (RT at 238:3-12.)',

        'Unsatisfied with this response, Ms. Vasquez submitted a formal written complaint '
        'to PCI\'s Ethics Hotline on May 8, 2024, detailing the discrepancies she had '
        'discovered. (CT at 95-97.) She also contacted the Department of Consumer Affairs '
        'on May 15, 2024 to report the suspected underreporting. (CT at 102.)',

        'Within weeks of her reports, Ms. Vasquez experienced a marked shift in her working '
        'conditions. On June 1, 2024, she was removed from the flagship product line and '
        'reassigned to a legacy system maintenance role. (CT at 108.) Her access to the QA '
        'database was revoked on June 15, 2024. (CT at 112.) On August 20, 2024, she received '
        'her first negative performance review in five years, citing "lack of teamwork" and '
        '"failure to follow chain of command." (CT at 118-120.)',
    ]

    for text in facts_paras:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    doc.add_page_break()

    # Argument
    h5 = doc.add_heading('ARGUMENT', level=1)
    h5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub1 = doc.add_heading(
        'I. THE TRIAL COURT ERRED IN GRANTING SUMMARY JUDGMENT ON THE '
        'WRONGFUL TERMINATION CLAIM', level=2
    )

    arg_paras = [
        'Summary judgment is properly granted only when "all the papers submitted show that '
        'there is no triable issue as to any material fact and that the moving party is entitled '
        'to judgment as a matter of law." (Code Civ. Proc., section 437c, subd. (c).) The court '
        'must view the evidence in the light most favorable to the nonmoving party. (Aguilar v. '
        'Atlantic Richfield Co. (2001) 25 Cal.4th 826, 843.)',

        'Here, substantial evidence raised triable issues of material fact as to whether '
        'Ms. Vasquez\'s termination was motivated by her protected whistleblowing activity. '
        'The temporal proximity between her reports in April-May 2024 and the adverse employment '
        'actions beginning in June 2024 is itself sufficient to create an inference of '
        'retaliatory motive. (See Loggins v. Kaiser Permanente International (2007) '
        '151 Cal.App.4th 1102, 1112-1113.)',

        'Moreover, the pretext evidence is compelling. Respondent claimed Ms. Vasquez was '
        'terminated for "restructuring," yet her position was filled by a new hire, '
        'David Morrison, within three weeks of her departure. (CT at 145.) The purported '
        'performance deficiencies cited in her termination letter were contradicted by five '
        'years of exemplary reviews. (CT at 48-52, 148.)',
    ]

    for text in arg_paras:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    # Conclusion
    doc.add_page_break()
    h6 = doc.add_heading('CONCLUSION', level=1)
    h6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    conc = doc.add_paragraph()
    conc.paragraph_format.first_line_indent = Inches(0.5)
    conc.paragraph_format.line_spacing = 2.0
    run = conc.add_run(
        'For the foregoing reasons, Appellant Elena Vasquez respectfully requests that this '
        'Court reverse the judgment of the trial court, vacate the order granting summary '
        'judgment, and remand for further proceedings consistent with this Court\'s opinion.'
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
