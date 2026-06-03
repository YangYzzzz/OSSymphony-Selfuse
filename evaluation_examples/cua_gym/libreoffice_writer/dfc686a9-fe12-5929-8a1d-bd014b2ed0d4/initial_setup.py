"""
Initial Setup: Create a realistic Final_Report.docx for the ExportToPDF macro task.
Task ID: writer_tm_088
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_088'
DOC_DIR = f'{WORKDIR}/Documents'
OUTPUT = f'{DOC_DIR}/Final_Report.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DOC_DIR, exist_ok=True)

    doc = Document()

    # --- Title ---
    title = doc.add_heading("Q4 2025 Financial Performance Report", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Prepared by the Finance Department")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)

    doc.add_paragraph()  # spacer

    # --- Executive Summary ---
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report presents the financial performance of Meridian Technologies Inc. "
        "for the fourth quarter of fiscal year 2025. Overall revenue increased by 12.3% "
        "compared to Q3, driven primarily by strong demand in the cloud services division. "
        "Operating margins improved to 23.7%, reflecting cost optimization initiatives "
        "implemented during the second half of the year."
    )

    # --- Revenue Breakdown ---
    doc.add_heading("Revenue Breakdown by Division", level=1)
    doc.add_paragraph(
        "The following table summarizes revenue contribution by each major division:"
    )

    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["Division", "Q4 Revenue ($M)", "Q3 Revenue ($M)", "Growth (%)"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ["Cloud Services", "$48.2", "$41.6", "15.9%"],
        ["Enterprise Software", "$32.7", "$30.1", "8.6%"],
        ["Consulting", "$18.4", "$17.9", "2.8%"],
        ["Hardware Solutions", "$12.1", "$11.3", "7.1%"],
        ["Total", "$111.4", "$100.9", "10.4%"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()

    # --- Key Highlights ---
    doc.add_heading("Key Highlights", level=1)
    highlights = [
        "Cloud Services division exceeded annual target by 8%, securing 14 new enterprise contracts.",
        "Customer retention rate improved to 94.2%, up from 91.8% in Q3.",
        "R&D expenditure increased to $15.3M, focused on AI-driven analytics platform.",
        "International revenue grew 18.7%, with strong performance in APAC and EMEA regions.",
        "Employee headcount reached 2,847, with 156 new hires in engineering roles.",
    ]
    for h in highlights:
        doc.add_paragraph(h, style="List Bullet")

    # --- Risk Factors ---
    doc.add_heading("Risk Factors", level=1)
    doc.add_paragraph(
        "Management has identified several risk factors that may impact future performance. "
        "Currency fluctuations in emerging markets continue to pose challenges for international "
        "revenue recognition. Additionally, increasing competition in the cloud infrastructure "
        "space from major technology providers could pressure margins in the medium term. "
        "The company has initiated hedging strategies and is investing in proprietary "
        "technology differentiation to mitigate these risks."
    )

    # --- Outlook ---
    doc.add_heading("Outlook for FY 2026", level=1)
    doc.add_paragraph(
        "Based on current pipeline analysis and market conditions, Meridian Technologies "
        "projects Q1 2026 revenue between $115M and $120M, representing approximately "
        "3-8% sequential growth. The board has approved a $25M capital expenditure plan "
        "for infrastructure expansion in the Dallas and Frankfurt data centers."
    )

    # --- Approval ---
    doc.add_paragraph()
    approval = doc.add_paragraph()
    approval.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = approval.add_run("Approved by: Victoria Chen, CFO")
    run.font.size = Pt(11)
    run.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = date_para.add_run("Date: December 15, 2025")
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
