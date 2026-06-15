"""
Initial Setup: Bulleted list with default spacing
Task ID: writer_lec_028
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_028'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title heading
    doc.add_heading("Q2 2025 Marketing Campaign Action Items", level=1)
    doc.add_paragraph(
        "The following action items were identified during the quarterly "
        "marketing review meeting held on March 28, 2025. Each team member "
        "is responsible for completing their assigned tasks before the next "
        "check-in on April 15."
    )

    # 10 bulleted list items with realistic content and default spacing
    items = [
        "Finalize the social media content calendar for April and May, including Instagram Reels and LinkedIn carousel posts",
        "Review and approve the updated brand guidelines document shared by the design team last Friday",
        "Schedule interviews with three customer success stories for the upcoming case study series",
        "Coordinate with the events team to confirm booth layout and promotional materials for the TechExpo conference",
        "Prepare a competitive analysis report comparing our Q1 engagement metrics against industry benchmarks",
        "Update the email nurture sequence for the enterprise segment with new product feature highlights",
        "Audit the current landing page conversion funnels and propose A/B testing variants for the pricing page",
        "Compile the monthly analytics dashboard with channel attribution data and share with senior leadership",
        "Draft the press release for the upcoming product launch, targeting a distribution date of April 22",
        "Organize a cross-functional brainstorming session with sales and product teams to align on Q3 messaging strategy",
    ]

    for item_text in items:
        doc.add_paragraph(item_text, style='List Bullet')

    # Add a closing paragraph
    doc.add_paragraph(
        "Please update the shared project tracker with your progress by end "
        "of day each Friday. Reach out to the project lead if you encounter "
        "any blockers or need additional resources."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
