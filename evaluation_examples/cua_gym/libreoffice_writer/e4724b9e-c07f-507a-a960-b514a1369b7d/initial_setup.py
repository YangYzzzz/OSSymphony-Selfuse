"""
Initial Setup: Insert BERT-Large benchmark scores from spreadsheet into research paper
Task ID: osworld_multi_apps_calc_to_writer_002
Domain: libreoffice_writer (multi-app: Writer + Calc)

Initial state:
  - research_paper.docx at /home/user/research_paper.docx
    (has Evaluation section — but NO table under it)
  - model_scores.xlsx at /home/user/Desktop/benchmarks/model_scores.xlsx
    (has Model/Accuracy/F1/Latency columns including BERT-Large row)
  - LibreOffice Writer opened with research_paper.docx
  - LibreOffice Calc opened with model_scores.xlsx
"""

import os
import shlex
import subprocess
import time
import openpyxl
from openpyxl.styles import Font, Alignment
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_multi_apps_calc_to_writer_002'
WRITER_OUTPUT = f'{WORKDIR}/research_paper.docx'
BENCHMARKS_DIR = f'{WORKDIR}/Desktop/benchmarks'
CALC_OUTPUT = f'{BENCHMARKS_DIR}/model_scores.xlsx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_model_scores_xlsx():
    """Create the benchmark scores spreadsheet with model evaluation data."""
    os.makedirs(BENCHMARKS_DIR, exist_ok=True)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Model Scores'

    # Headers
    headers = ['Model', 'Accuracy', 'F1', 'Latency (ms)']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(name='Calibri', bold=True, size=11)
        cell.alignment = Alignment(horizontal='center')

    # Data — realistic NLP benchmark scores for various models
    data = [
        ['GPT-4',           0.9312, 0.9278, 145.3],
        ['GPT-3.5-Turbo',   0.8874, 0.8801,  78.6],
        ['BERT-Large',      0.9145, 0.9089,  62.4],
        ['BERT-Base',       0.8823, 0.8752,  38.2],
        ['RoBERTa-Large',   0.9201, 0.9163,  68.7],
        ['RoBERTa-Base',    0.8911, 0.8845,  41.5],
        ['DistilBERT',      0.8634, 0.8571,  22.1],
        ['ALBERT-XXLarge',  0.9267, 0.9218,  95.3],
        ['ALBERT-Base',     0.8712, 0.8649,  31.8],
        ['DeBERTa-v3',      0.9388, 0.9342, 110.5],
        ['XLNet-Large',     0.9156, 0.9098,  88.9],
        ['T5-Large',        0.9023, 0.8967, 134.2],
    ]
    for r, row_data in enumerate(data, 2):
        for c, val in enumerate(row_data, 1):
            ws.cell(row=r, column=c, value=val)

    # Set column widths for readability
    ws.column_dimensions['A'].width = 18
    ws.column_dimensions['B'].width = 12
    ws.column_dimensions['C'].width = 10
    ws.column_dimensions['D'].width = 16

    wb.save(CALC_OUTPUT)
    print(f'Benchmark spreadsheet created: {CALC_OUTPUT}')


def create_research_paper_docx():
    """Create the research paper document with an Evaluation section (no table yet)."""
    doc = Document()

    # Title
    title = doc.add_heading('Neural Language Model Evaluation: A Comparative Study', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Authors
    authors_para = doc.add_paragraph('Authors: Dr. Emily Watkins, Prof. James Liu, Anika Sharma')
    authors_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')  # blank line

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This paper presents a systematic evaluation of modern pre-trained language models '
        'on a suite of natural language understanding benchmarks. We compare architectures '
        'including encoder-only models (BERT, RoBERTa, ALBERT, DeBERTa), decoder-only '
        'models (GPT-3.5, GPT-4), and encoder-decoder models (T5). Our evaluation covers '
        'accuracy, F1 score, and inference latency, providing practitioners with actionable '
        'guidance for model selection in production environments.'
    )

    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'The proliferation of large pre-trained language models has transformed the landscape '
        'of natural language processing. Since the seminal introduction of BERT (Devlin et al., '
        '2019), dozens of architectural variants have emerged, each claiming improvements over '
        'prior work. Practitioners face a growing challenge: selecting the right model for '
        'a given application requires balancing performance metrics against latency, memory '
        'footprint, and fine-tuning cost.'
    )
    doc.add_paragraph(
        'In this work, we standardize the evaluation protocol across twelve widely-used '
        'models and report results on the GLUE and SuperGLUE benchmark suites. We pay '
        'particular attention to inference latency, which is often underreported but critical '
        'in real-time applications such as conversational agents and information retrieval systems.'
    )

    # 2. Related Work
    doc.add_heading('2. Related Work', level=1)
    doc.add_paragraph(
        'Comparative studies of language models have been conducted by several groups. '
        'Wang et al. (2021) evaluated BERT-family models across sentiment analysis tasks. '
        'Hernandez and Brown (2022) focused on efficiency-performance trade-offs for '
        'edge deployment scenarios. Our study extends this line of research by including '
        'more recent architectures and providing a unified latency measurement methodology.'
    )

    # 3. Methodology
    doc.add_heading('3. Methodology', level=1)
    doc.add_paragraph(
        'All models were evaluated using identical preprocessing pipelines and hardware '
        'configurations. Accuracy and F1 scores were measured on the development split of '
        'the MNLI dataset (Williams et al., 2018). Latency was measured as the median '
        'inference time per sample over 1,000 forward passes on an NVIDIA A100 GPU with '
        'batch size 1 to simulate single-query production scenarios.'
    )
    doc.add_paragraph(
        'Models were loaded with their default tokenizers from the HuggingFace Transformers '
        'library (version 4.35). For each model, we report the zero-shot performance without '
        'any task-specific fine-tuning to ensure fair comparison across architectures.'
    )

    # 4. Evaluation  — key section: NO table here in initial state
    doc.add_heading('Evaluation', level=1)
    doc.add_paragraph(
        'We evaluated all twelve models on the full benchmark suite. The results demonstrate '
        'clear trade-offs between model size, accuracy, and inference speed. Larger models '
        'consistently achieve higher accuracy and F1 scores at the cost of increased latency. '
        'Notably, DeBERTa-v3 achieves the highest accuracy (93.88%) among all evaluated models, '
        'while DistilBERT offers the lowest latency (22.1 ms) with only modest performance degradation.'
    )
    doc.add_paragraph(
        'Among the BERT family, BERT-Large serves as the canonical reference point. Its '
        'balance of performance and latency has made it the de facto baseline for many '
        'downstream NLP tasks. The full evaluation scores for individual models are '
        'summarized in the benchmark results table.'
    )

    # 5. Discussion
    doc.add_heading('5. Discussion', level=1)
    doc.add_paragraph(
        'The results highlight several practical considerations for model deployment. '
        'First, the accuracy gap between base and large variants of the same architecture '
        'is generally modest (1-3%), while the latency increase can be substantial (60-80%). '
        'This suggests that base models may be preferable for latency-sensitive applications '
        'with minimal accuracy loss.'
    )
    doc.add_paragraph(
        'Second, the GPT-4 model achieves competitive accuracy but with significantly higher '
        'latency compared to encoder-only models, making it less suitable for real-time '
        'classification tasks. However, its general-purpose capabilities may justify the '
        'cost in applications requiring diverse language understanding.'
    )

    # 6. Conclusion
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        'This paper provides a comprehensive evaluation of twelve modern language models, '
        'covering performance and latency dimensions. Our findings support the use of '
        'DeBERTa-v3 for accuracy-critical tasks and DistilBERT for latency-constrained '
        'scenarios. Future work will extend this evaluation to multilingual benchmarks '
        'and examine model calibration properties.'
    )

    # References
    doc.add_heading('References', level=1)
    doc.add_paragraph(
        'Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of '
        'Deep Bidirectional Transformers for Language Understanding. NAACL-HLT 2019.'
    )
    doc.add_paragraph(
        'Wang, A., et al. (2019). GLUE: A Multi-Task Benchmark and Analysis Platform for '
        'Natural Language Understanding. ICLR 2019.'
    )
    doc.add_paragraph(
        'Hernandez, D., & Brown, T. (2022). Scaling Laws for Transfer. arXiv:2102.01293.'
    )

    doc.save(WRITER_OUTPUT)
    print(f'Research paper created: {WRITER_OUTPUT}')


def create_initial():
    create_model_scores_xlsx()
    create_research_paper_docx()

    # GUI-ready startup: open LibreOffice Writer with the research paper,
    # then LibreOffice Calc with the benchmark spreadsheet
    launch_gui(f'libreoffice --writer "{WRITER_OUTPUT}"', delay_sec=2.5)
    launch_gui(f'libreoffice --calc "{CALC_OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer (research_paper.docx) and Calc (model_scores.xlsx) with DISPLAY=:0')


create_initial()
