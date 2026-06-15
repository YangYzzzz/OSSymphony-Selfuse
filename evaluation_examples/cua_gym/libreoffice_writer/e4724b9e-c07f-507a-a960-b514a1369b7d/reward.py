"""
Reward Script: Extract BERT-Large scores from xlsx and insert as table into Evaluation section
Task ID: osworld_multi_apps_calc_to_writer_002
Domain: libreoffice_writer (multi-app: calc + writer)
Scoring:
  Component 1: A table exists in the document (0.3 pts)
  Component 2: Table has correct column headers — Model, Accuracy, F1, Latency (0.3 pts)
  Component 3: Table contains BERT-Large row with correct values (Acc=0.9145, F1=0.9089, Latency=62.4) (0.4 pts)
Total: 1.0
"""

import os

# Domain-specific import
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_calc_to_writer_002'

# BERT-Large ground truth values (from model_scores.xlsx)
BERT_ACCURACY = 0.9145
BERT_F1 = 0.9089
BERT_LATENCY = 62.4
NUMERIC_TOLERANCE = 0.001


def parse_float(val):
    """Safely parse a cell value as float."""
    if val is None:
        return None
    try:
        return float(str(val).strip())
    except (ValueError, TypeError):
        return None


def find_bert_large_row(table):
    """
    Search table rows for BERT-Large entry.
    Returns (accuracy, f1, latency) tuple or None if not found.
    """
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if cells and "bert-large" in cells[0].lower() and len(cells) >= 4:
            acc = parse_float(cells[1])
            f1 = parse_float(cells[2])
            lat = parse_float(cells[3])
            return (acc, f1, lat)
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task requires inserting a table with BERT-Large's benchmark scores
    (Accuracy, F1, Latency) into the 'Evaluation' section of the research paper.

    Initial state: 0 tables in document
    Golden state:  1 table with header row and BERT-Large data row
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ----- Component 1: A table exists in the document (0.3 points) -----
    # Initial state has 0 tables; golden state has 1 table.
    # This is the primary task-introduced change.
    try:
        num_tables = len(doc.tables)
        if num_tables >= 1:
            print(f"PASS: Component 1 — document contains {num_tables} table(s) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — expected at least 1 table, found {num_tables}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ----- Component 2: Table has correct column headers (0.3 points) -----
    # Headers expected: Model, Accuracy, F1, Latency (case-insensitive partial match)
    try:
        if len(doc.tables) >= 1:
            header_row = doc.tables[0].rows[0]
            cell_texts = [cell.text.strip() for cell in header_row.cells]
            cell_lower = [t.lower() for t in cell_texts]

            headers_ok = (
                any("model" in t for t in cell_lower) and
                any("accuracy" in t for t in cell_lower) and
                any("f1" in t for t in cell_lower) and
                any("latency" in t for t in cell_lower)
            )

            if headers_ok:
                print(f"PASS: Component 2 — table headers correct: {cell_texts} (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — expected headers [Model, Accuracy, F1, Latency], found: {cell_texts}")
        else:
            print("FAIL: Component 2 — no table found, cannot check headers")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ----- Component 3: Table contains BERT-Large row with correct values (0.4 points) -----
    # BERT-Large ground truth: Accuracy=0.9145, F1=0.9089, Latency=62.4
    # The row must reference 'BERT-Large' (case-insensitive) and have numeric values
    # within tolerance of the expected values from model_scores.xlsx.
    try:
        if len(doc.tables) >= 1:
            result = find_bert_large_row(doc.tables[0])
            if result is None:
                print("FAIL: Component 3 — no row containing 'BERT-Large' found in table")
            else:
                acc_val, f1_val, lat_val = result
                print(f"INFO: BERT-Large row found — Accuracy={acc_val} (expected {BERT_ACCURACY}), "
                      f"F1={f1_val} (expected {BERT_F1}), Latency={lat_val} (expected {BERT_LATENCY})")

                values_ok = (
                    acc_val is not None and abs(acc_val - BERT_ACCURACY) <= NUMERIC_TOLERANCE and
                    f1_val is not None and abs(f1_val - BERT_F1) <= NUMERIC_TOLERANCE and
                    lat_val is not None and abs(lat_val - BERT_LATENCY) <= NUMERIC_TOLERANCE
                )

                if values_ok:
                    print(f"PASS: Component 3 — BERT-Large row with correct values found (0.4 pts)")
                    total_score += 0.4
                else:
                    print(f"FAIL: Component 3 — BERT-Large row found but values incorrect: "
                          f"acc_ok={acc_val is not None and abs(acc_val - BERT_ACCURACY) <= NUMERIC_TOLERANCE}, "
                          f"f1_ok={f1_val is not None and abs(f1_val - BERT_F1) <= NUMERIC_TOLERANCE}, "
                          f"latency_ok={lat_val is not None and abs(lat_val - BERT_LATENCY) <= NUMERIC_TOLERANCE}")
        else:
            print("FAIL: Component 3 — no table found, cannot check BERT-Large row")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/research_paper.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
