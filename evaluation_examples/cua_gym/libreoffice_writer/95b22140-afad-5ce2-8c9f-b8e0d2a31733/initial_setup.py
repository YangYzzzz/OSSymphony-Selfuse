"""
Initial Setup: NDA template with unprotected 'Confidential Terms' section
Task ID: writer_struct_020
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_020'
OUTPUT = f'{WORKDIR}/Desktop/nda_template.docx'

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_sdt_section(section_name: str, paragraphs_xml: list, protected: bool = False):
    """
    Create a LibreOffice-style named section as a <w:sdt> element.
    paragraphs_xml: list of lxml elements (w:p) to include in the section body.
    protected: whether to add content locking (sdtContentLocked).
    """
    sdt = etree.Element(f'{{{W}}}sdt')

    # sdtPr
    sdtPr = etree.SubElement(sdt, f'{{{W}}}sdtPr')
    alias = etree.SubElement(sdtPr, f'{{{W}}}alias')
    alias.set(f'{{{W}}}val', section_name)
    tag = etree.SubElement(sdtPr, f'{{{W}}}tag')
    tag.set(f'{{{W}}}val', section_name)

    if protected:
        lock = etree.SubElement(sdtPr, f'{{{W}}}lock')
        lock.set(f'{{{W}}}val', 'sdtContentLocked')

    # sdtContent
    sdtContent = etree.SubElement(sdt, f'{{{W}}}sdtContent')
    for para_elem in paragraphs_xml:
        sdtContent.append(para_elem)

    return sdt


def make_paragraph_elem(text: str, font_name: str = 'Times New Roman', font_size_pt: int = 11,
                         space_after_pt: int = 6, bold: bool = False) -> etree._Element:
    """Create a w:p element with text."""
    p = etree.Element(f'{{{W}}}p')
    pPr = etree.SubElement(p, f'{{{W}}}pPr')
    spacing = etree.SubElement(pPr, f'{{{W}}}spacing')
    spacing.set(f'{{{W}}}after', str(space_after_pt * 20))  # twips
    r = etree.SubElement(p, f'{{{W}}}r')
    rPr = etree.SubElement(r, f'{{{W}}}rPr')
    sz = etree.SubElement(rPr, f'{{{W}}}sz')
    sz.set(f'{{{W}}}val', str(font_size_pt * 2))  # half-points
    szCs = etree.SubElement(rPr, f'{{{W}}}szCs')
    szCs.set(f'{{{W}}}val', str(font_size_pt * 2))
    rFonts = etree.SubElement(rPr, f'{{{W}}}rFonts')
    rFonts.set(f'{{{W}}}ascii', font_name)
    rFonts.set(f'{{{W}}}hAnsi', font_name)
    if bold:
        b = etree.SubElement(rPr, f'{{{W}}}b')
    t = etree.SubElement(r, f'{{{W}}}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    return p


def create_initial():
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Reasonable page margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ===== PAGE 1: Header and Parties =====
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(6)
    run = title_para.add_run('NON-DISCLOSURE AGREEMENT')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    sub_para = doc.add_paragraph()
    sub_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_para.paragraph_format.space_after = Pt(12)
    run = sub_para.add_run('(Mutual Non-Disclosure and Confidentiality Agreement)')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.italic = True

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    run = intro.add_run(
        'This Non-Disclosure Agreement ("Agreement") is entered into as of March 1, 2025, '
        'by and between:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    p1 = doc.add_paragraph()
    p1.paragraph_format.left_indent = Inches(0.5)
    p1.paragraph_format.space_after = Pt(6)
    run = p1.add_run('Nexus Innovations, LLC')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run2 = p1.add_run(
        ', a limited liability company organized under the laws of the State of Delaware, '
        'with its principal place of business at 400 Market Street, Suite 1200, '
        'Wilmington, DE 19801 ("Disclosing Party"); and'
    )
    run2.font.size = Pt(11)
    run2.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.5)
    p2.paragraph_format.space_after = Pt(12)
    run = p2.add_run('Horizon Analytics, Inc.')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run2 = p2.add_run(
        ', a corporation incorporated under the laws of the State of California, '
        'with its principal place of business at 2250 Innovation Drive, Suite 800, '
        'San Jose, CA 95131 ("Receiving Party").'
    )
    run2.font.size = Pt(11)
    run2.font.name = 'Times New Roman'

    h1 = doc.add_heading('1. Purpose', level=1)
    for run in h1.runs:
        run.font.name = 'Times New Roman'

    p_purpose = doc.add_paragraph()
    p_purpose.paragraph_format.space_after = Pt(6)
    run = p_purpose.add_run(
        'The parties wish to explore a potential business relationship concerning the '
        'joint development of advanced data analytics software and related intellectual '
        'property ("Business Purpose"). In connection with this exploration, each party '
        'may disclose to the other certain confidential and proprietary information. '
        'The parties intend this Agreement to protect such information.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    h2 = doc.add_heading('2. Definition of Confidential Information', level=1)
    for run in h2.runs:
        run.font.name = 'Times New Roman'

    p_def = doc.add_paragraph()
    p_def.paragraph_format.space_after = Pt(6)
    run = p_def.add_run(
        '"Confidential Information" means any and all information or data that has or '
        'could have commercial value or other utility in the business in which the '
        'Disclosing Party is engaged. Confidential Information includes, without limitation: '
        'trade secrets, inventions, ideas, processes, computer source and object code, '
        'formulas, data, programs, software, other works of authorship, know-how, improvements, '
        'discoveries, developments, designs, techniques, marketing plans, client lists, '
        'financial data, and business strategies.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # Page break to page 2
    doc.add_page_break()

    # ===== PAGE 2: Obligations =====
    h3 = doc.add_heading('3. Obligations of Receiving Party', level=1)
    for run in h3.runs:
        run.font.name = 'Times New Roman'

    p_oblig = doc.add_paragraph()
    p_oblig.paragraph_format.space_after = Pt(6)
    run = p_oblig.add_run(
        'The Receiving Party agrees to: (a) hold the Confidential Information in strict confidence; '
        '(b) not to disclose the Confidential Information to any third parties without the prior '
        'written consent of the Disclosing Party; (c) use the Confidential Information solely for '
        'the Business Purpose; and (d) protect the Confidential Information using the same degree '
        'of care it uses to protect its own confidential information, but in no event less than '
        'reasonable care.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    h4 = doc.add_heading('4. Term', level=1)
    for run in h4.runs:
        run.font.name = 'Times New Roman'

    p_term = doc.add_paragraph()
    p_term.paragraph_format.space_after = Pt(6)
    run = p_term.add_run(
        'This Agreement shall commence on the date first written above and shall remain in full '
        'force and effect for a period of three (3) years, unless earlier terminated by either '
        'party upon thirty (30) days written notice to the other party. Obligations with respect '
        'to Confidential Information shall survive the termination of this Agreement for an '
        'additional period of five (5) years.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    h5 = doc.add_heading('5. Exceptions', level=1)
    for run in h5.runs:
        run.font.name = 'Times New Roman'

    p_except = doc.add_paragraph()
    p_except.paragraph_format.space_after = Pt(6)
    run = p_except.add_run(
        'The obligations set forth in this Agreement shall not apply to information that: '
        '(a) was publicly known at the time of disclosure; (b) becomes publicly known through '
        'no fault of the Receiving Party; (c) was rightfully received from a third party without '
        'restriction on disclosure; or (d) was independently developed by the Receiving Party '
        'without use of or reference to the Confidential Information.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # ===== Section 6 heading (outside SDT) =====
    h_conf = doc.add_heading('6. Confidential Terms', level=1)
    for run in h_conf.runs:
        run.font.name = 'Times New Roman'

    # ===== "Confidential Terms" Named Section via SDT (UNPROTECTED for initial) =====
    # Build the two paragraphs for this named section
    text1 = (
        'The parties acknowledge that all information exchanged under this Agreement constitutes '
        'highly sensitive proprietary data. The Receiving Party shall implement and maintain '
        'comprehensive technical and organizational security measures to protect such information, '
        'including but not limited to: encrypted storage systems, role-based access controls, '
        'audit logging of all access events, and mandatory security training for all personnel '
        'who handle Confidential Information.'
    )
    text2 = (
        'Any breach of the confidentiality obligations set forth in this section shall entitle '
        'the Disclosing Party to seek immediate injunctive relief without bond, in addition to '
        'any other remedies available at law or in equity. The parties agree that monetary damages '
        'alone would be insufficient to compensate for any such breach, and that specific '
        'performance or injunctive relief shall be available as remedies.'
    )

    p1_elem = make_paragraph_elem(text1)
    p2_elem = make_paragraph_elem(text2)

    # Create SDT for named section - UNPROTECTED (no lock element)
    sdt = make_sdt_section('Confidential Terms', [p1_elem, p2_elem], protected=False)

    # Insert SDT before the final sectPr in body
    body = doc.element.body
    last_elem = body[-1]
    body.insert(list(body).index(last_elem), sdt)

    # Page break before page 3
    doc.add_page_break()

    # ===== PAGE 3: General Provisions =====
    h6 = doc.add_heading('7. General Provisions', level=1)
    for run in h6.runs:
        run.font.name = 'Times New Roman'

    for text in [
        '7.1 Governing Law. This Agreement shall be governed by and construed in accordance '
        'with the laws of the State of Delaware, without regard to its conflict of laws provisions.',
        '7.2 Entire Agreement. This Agreement constitutes the entire agreement between the parties '
        'with respect to its subject matter and supersedes all prior negotiations, representations, '
        'warranties, and understandings of the parties with respect thereto.',
        '7.3 Amendments. No amendment or modification of this Agreement shall be valid unless '
        'made in writing and duly executed by authorized representatives of both parties.',
        '7.4 Severability. If any provision of this Agreement is found to be unenforceable, '
        'the remainder shall be enforced as fully as possible, and the unenforceable provision '
        'shall be deemed modified to the limited extent required to permit its enforcement in a '
        'manner most closely representing the parties\' original intention.',
        '7.5 Counterparts. This Agreement may be executed in counterparts, each of which shall '
        'be deemed an original, and all of which together shall constitute one and the same instrument.',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = 'Times New Roman'

    sig_intro = doc.add_paragraph()
    sig_intro.paragraph_format.space_before = Pt(12)
    sig_intro.paragraph_format.space_after = Pt(12)
    run = sig_intro.add_run(
        'IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    for label, company in [
        ('Nexus Innovations, LLC', True),
        ('Horizon Analytics, Inc.', True),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = 'Times New Roman'
        for line_label in ['Signature', 'Name', 'Title', 'Date']:
            lp = doc.add_paragraph()
            lp.paragraph_format.space_after = Pt(4)
            lr = lp.add_run(f'{line_label}: _________________________________')
            lr.font.size = Pt(11)
            lr.font.name = 'Times New Roman'
        doc.add_paragraph()  # spacer

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
