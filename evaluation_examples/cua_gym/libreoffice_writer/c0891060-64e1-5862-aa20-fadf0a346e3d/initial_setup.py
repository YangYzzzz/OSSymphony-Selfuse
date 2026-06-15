"""
Initial Setup: Thesis abstract document with default formatting
Task ID: writer_acad_078
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_078'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default page margins for a thesis-like document
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title: "Abstract" in plain Normal style (NOT centered, NOT Heading 1)
    title_para = doc.add_paragraph("Abstract")
    # Leave as default left-aligned Normal style - no centering, no heading style

    # Abstract body text - realistic thesis abstract about renewable energy
    abstract_text = (
        "The increasing urgency of climate change mitigation has driven significant "
        "investment in renewable energy technologies across both developed and developing "
        "nations. This thesis examines the socioeconomic impacts of transitioning from "
        "fossil fuel-dependent energy systems to diversified renewable portfolios in "
        "three Southeast Asian economies: Vietnam, Thailand, and the Philippines. "
        "Using a mixed-methods approach combining econometric modeling with qualitative "
        "case studies conducted between 2022 and 2024, we analyze employment shifts, "
        "household energy expenditure patterns, and regional GDP contributions associated "
        "with solar photovoltaic, onshore wind, and small-scale hydroelectric installations. "
        "Our quantitative findings indicate that renewable energy deployment correlates "
        "with a net increase of 2.3 jobs per megawatt of installed capacity in rural "
        "regions, while urban areas experience a modest displacement effect of 0.7 jobs "
        "per megawatt in conventional power generation. Household energy costs decreased "
        "by an average of 14.6% in communities within 50 kilometers of new renewable "
        "installations, with the most pronounced savings observed in off-grid island "
        "communities previously reliant on diesel generators. The qualitative analysis "
        "reveals that community acceptance is strongly mediated by participatory planning "
        "processes and transparent benefit-sharing mechanisms. Policy implications suggest "
        "that targeted workforce retraining programs, coupled with community ownership "
        "models, can maximize the socioeconomic co-benefits of energy transition while "
        "minimizing distributional inequities across income quintiles."
    )
    body_para = doc.add_paragraph(abstract_text)
    # Default formatting - no explicit single-spacing set

    # Keywords line - NOT in italics (task requires making it italic)
    keywords_para = doc.add_paragraph()
    keywords_run = keywords_para.add_run(
        "Keywords: renewable energy, socioeconomic impact, energy transition, "
        "Southeast Asia, employment, household expenditure, community acceptance"
    )
    # Leave as normal (non-italic) text

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
