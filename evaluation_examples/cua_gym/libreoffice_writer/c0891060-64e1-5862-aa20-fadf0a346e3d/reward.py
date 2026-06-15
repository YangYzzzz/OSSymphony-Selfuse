"""
Reward Script: Format thesis abstract with centered Heading 1 title,
single-spaced abstract paragraph, and italicized keywords line.
Task ID: writer_acad_078
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35) - 'Abstract' title is Heading 1 style and centered
  Component 2 (0.35) - Abstract body paragraph is single-spaced (line_spacing == 1.0)
  Component 3 (0.30) - Keywords paragraph has all runs in italic
"""

import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_078'


def persist_app_state(domain: str):
    """Attempt to save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 3 paragraphs
    if len(doc.paragraphs) < 3:
        print(f"CRITICAL: Expected at least 3 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Identify paragraphs by content
    title_para = None
    body_para = None
    keywords_para = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if text.lower() == 'abstract':
            title_para = para
        elif text.lower().startswith('keywords:') or text.lower().startswith('keywords :'):
            keywords_para = para
        elif len(text) > 100:
            # The abstract body is the long paragraph
            body_para = para

    if title_para is None:
        print("CRITICAL: Could not find 'Abstract' title paragraph")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: 'Abstract' is Heading 1 style AND centered (0.35 points)
    # Initial: Normal style, no alignment -> should FAIL
    # Golden: Heading 1, CENTER -> should PASS
    try:
        style_name = title_para.style.name if title_para.style else None
        alignment = title_para.paragraph_format.alignment

        is_heading = style_name is not None and 'heading' in style_name.lower() and '1' in style_name
        is_centered = alignment == WD_PARAGRAPH_ALIGNMENT.CENTER

        if is_heading and is_centered:
            print(f"PASS: Component 1 - 'Abstract' is {style_name} and centered (0.35 pts)")
            total_score += 0.35
        elif is_heading:
            print(f"PARTIAL: Component 1 - 'Abstract' is {style_name} but not centered (0.15 pts)")
            total_score += 0.15
        elif is_centered:
            print(f"PARTIAL: Component 1 - 'Abstract' is centered but style is {style_name} (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 - style={style_name}, alignment={alignment}")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Abstract body paragraph is single-spaced (0.35 points)
    # Initial: line_spacing is None (default) -> should FAIL
    # Golden: line_spacing is 1.0 -> should PASS
    try:
        if body_para is None:
            print("FAIL: Component 2 - Could not find abstract body paragraph")
        else:
            line_spacing = body_para.paragraph_format.line_spacing
            # Single-spaced means line_spacing == 1.0
            if line_spacing is not None and abs(float(line_spacing) - 1.0) < 0.05:
                print(f"PASS: Component 2 - Abstract body is single-spaced (line_spacing={line_spacing}) (0.35 pts)")
                total_score += 0.35
            else:
                print(f"FAIL: Component 2 - Expected line_spacing=1.0, found {line_spacing}")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Keywords paragraph has all runs in italic (0.30 points)
    # Initial: italic is None -> should FAIL
    # Golden: italic is True -> should PASS
    try:
        if keywords_para is None:
            print("FAIL: Component 3 - Could not find keywords paragraph")
        else:
            runs = keywords_para.runs
            if len(runs) == 0:
                print("FAIL: Component 3 - Keywords paragraph has no runs")
            else:
                all_italic = all(
                    run.font.italic is True
                    for run in runs
                    if run.text.strip()  # skip whitespace-only runs
                )
                if all_italic:
                    print(f"PASS: Component 3 - Keywords paragraph is fully italic ({len(runs)} runs) (0.30 pts)")
                    total_score += 0.30
                else:
                    non_italic = [
                        (r.text[:30], r.font.italic)
                        for r in runs if r.text.strip() and r.font.italic is not True
                    ]
                    print(f"FAIL: Component 3 - Not all runs italic. Non-italic: {non_italic}")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
