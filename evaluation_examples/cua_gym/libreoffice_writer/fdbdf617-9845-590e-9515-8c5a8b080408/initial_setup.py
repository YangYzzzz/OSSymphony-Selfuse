"""
Initial Setup: Insert cross-reference to 'Chapter 4: Discussion' heading
Task ID: writer_struct_017
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_017'
OUTPUT = f'{WORKDIR}/Desktop/phd_thesis.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_with_bookmark(doc, text, level, bookmark_name):
    """Add a Heading paragraph and attach a bookmark to it."""
    heading = doc.add_heading(text, level=level)
    # Insert bookmark start and end around the heading run
    heading_run = heading.runs[0] if heading.runs else heading.add_run(text)

    # Create bookmark start element
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(abs(hash(bookmark_name)) % 100000))
    bm_start.set(qn('w:name'), bookmark_name)

    # Create bookmark end element
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(abs(hash(bookmark_name)) % 100000))

    # Insert bookmark around run in paragraph XML
    run_elem = heading_run._element
    run_elem.addprevious(bm_start)
    run_elem.addnext(bm_end)

    return heading


def create_initial():
    # Ensure Desktop directory exists
    desktop_dir = f'{WORKDIR}/Desktop'
    os.makedirs(desktop_dir, exist_ok=True)

    doc = Document()

    # Set default font and margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ===== Title Page =====
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = 1  # CENTER
    title_para.paragraph_format.space_before = Pt(72)
    run = title_para.add_run('Longitudinal Analysis of Cognitive Load in\nAdaptive Learning Environments')
    run.bold = True
    run.font.size = Pt(18)

    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = 1
    author_para.paragraph_format.space_before = Pt(48)
    run = author_para.add_run('Dr. Emily Hartwell\nDepartment of Educational Psychology\nUniversity of Westfield\n2024')
    run.font.size = Pt(14)

    doc.add_page_break()

    # ===== Abstract =====
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This dissertation investigates the dynamic relationship between cognitive load and adaptive '
        'learning systems over a period of three academic semesters. Using a mixed-methods design '
        'combining eye-tracking, EEG measurements, and structured interviews, we examined 142 '
        'undergraduate students enrolled in adaptive online courses at two research universities. '
        'Our findings suggest that effective scaffolding reduces extraneous cognitive load by 34% '
        'while maintaining germane load levels critical to schema formation. Theoretical and '
        'practical implications for instructional design are discussed in detail.'
    )
    doc.add_paragraph(
        'Keywords: cognitive load theory, adaptive learning, e-learning, scaffolding, '
        'instructional design, educational technology'
    )

    doc.add_page_break()

    # ===== Table of Contents =====
    doc.add_heading('Table of Contents', level=1)
    toc_entries = [
        ('Chapter 1: Introduction', '4'),
        ('Chapter 2: Literature Review', '7'),
        ('Chapter 3: Methodology', '11'),
        ('Chapter 4: Discussion', '15'),
        ('Chapter 5: Conclusion', '18'),
        ('References', '20'),
    ]
    for entry_text, page_num in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(entry_text)
        p.add_run(f'\t{page_num}')

    doc.add_page_break()

    # ===== Chapter 1: Introduction =====
    add_heading_with_bookmark(doc, 'Chapter 1: Introduction', 1, '_Chapter1_Introduction')

    doc.add_paragraph(
        'The proliferation of digital learning platforms over the past decade has generated '
        'unprecedented opportunities for personalized education. Adaptive learning systems '
        '(ALS) represent a paradigm shift from static curricula toward dynamically responsive '
        'instructional pathways calibrated to individual learner profiles. Despite growing '
        'adoption in higher education contexts, fundamental questions remain regarding the '
        'cognitive mechanisms through which such systems influence learning outcomes.'
    )
    doc.add_paragraph(
        'Cognitive Load Theory (CLT), originally proposed by Sweller (1988) and subsequently '
        'expanded by Paas and van Merriënboer (1994), provides a compelling theoretical '
        'framework for understanding the interaction between instructional design and working '
        'memory constraints. The theory distinguishes three components of cognitive load: '
        'intrinsic, extraneous, and germane. Intrinsic load is inherent to the complexity of '
        'the learning material; extraneous load arises from poorly designed instruction; '
        'germane load reflects cognitive resources devoted to schema construction.'
    )
    doc.add_paragraph(
        'Central to this investigation is the hypothesis that adaptive feedback mechanisms '
        'within digital learning environments can modulate the distribution of cognitive load '
        'across these three components. Specifically, we propose that timely, context-sensitive '
        'scaffolding reduces extraneous load while preserving the germane load necessary for '
        'deep conceptual understanding.'
    )
    doc.add_paragraph(
        'This research draws upon longitudinal data collected across three consecutive semesters, '
        'encompassing both quantitative physiological measures and qualitative learner accounts. '
        'The study contributes to a growing body of empirical work on cognitive load in '
        'technology-mediated environments, with particular attention to the role of adaptive '
        'feedback timing and modality. Broader implications for the design of next-generation '
        'adaptive systems are addressed throughout the analysis.'
    )
    doc.add_paragraph(
        '1.1 Research Questions\n'
        'This dissertation is guided by three primary research questions:\n'
        '(1) How does the deployment of adaptive scaffolding affect intrinsic, extraneous, and '
        'germane cognitive load in undergraduate learners?\n'
        '(2) Are there measurable differences in cognitive load profiles between high-performing '
        'and low-performing student cohorts when exposed to adaptive versus static instruction?\n'
        '(3) What instructional design features of adaptive learning systems are most predictive '
        'of reduced extraneous cognitive load over time?'
    )
    doc.add_paragraph(
        '1.2 Significance of the Study\n'
        'The present study addresses a critical gap in empirical research on adaptive learning '
        'by integrating multi-modal physiological measurement with learner self-reports and '
        'academic performance data. Prior studies have predominantly relied on single-measure '
        'proxies for cognitive load, limiting the ecological validity of their conclusions. '
        'By triangulating across measurement modalities, we aim to provide a more robust '
        'characterization of cognitive load dynamics in naturalistic learning environments.'
    )

    doc.add_page_break()

    # ===== Chapter 2: Literature Review =====
    add_heading_with_bookmark(doc, 'Chapter 2: Literature Review', 1, '_Chapter2_LiteratureReview')

    doc.add_paragraph(
        'The literature on cognitive load in technology-enhanced learning environments has expanded '
        'considerably since the formalization of CLT. Early empirical studies primarily employed '
        'subjective rating scales derived from the work of Paas (1992), using single-item '
        'nine-point scales administered immediately following learning tasks. These approaches, '
        'while pragmatically convenient, have been critiqued for their susceptibility to '
        'retrospective bias and inability to capture within-task load fluctuations.'
    )
    doc.add_paragraph(
        '2.1 Adaptive Learning Systems: Historical Development\n'
        'The conceptual origins of adaptive instruction trace to early work in programmed learning '
        'and intelligent tutoring systems (ITS). Carbonell\'s (1970) SCHOLAR system represented '
        'a foundational attempt at adapting instructional sequences to learner knowledge states. '
        'Subsequent developments in the 1980s and 1990s, including the ACT-R based Cognitive '
        'Tutor (Anderson et al., 1995), demonstrated statistically significant learning gains '
        'over conventional instruction in controlled studies.'
    )
    doc.add_paragraph(
        '2.2 Physiological Correlates of Cognitive Load\n'
        'Advances in wearable sensor technology have enabled increasingly fine-grained '
        'measurement of cognitive load during authentic learning activities. Pupillometry—the '
        'measurement of pupil dilation as a proxy for mental effort—has demonstrated robust '
        'correlations with working memory demand across multiple task domains (Beatty, 1982; '
        'Kahneman & Beatty, 1966). More recently, electroencephalography (EEG) has been '
        'applied to derive real-time indices of cognitive load, with theta-band power at frontal '
        'electrodes reliably increasing with task difficulty (Klimesch, 1999).'
    )
    doc.add_paragraph(
        '2.3 Scaffolding and Load Reduction\n'
        'Scaffolding, as conceptualized by Wood, Bruner, and Ross (1976), refers to temporary '
        'support structures that enable learners to engage with tasks beyond their current '
        'independent capacity. In digital contexts, scaffolding may take the form of worked '
        'examples, process prompts, adaptive hints, or metacognitive feedback. Research by '
        'Sweller and Cooper (1985) demonstrated that studying worked examples could dramatically '
        'reduce extraneous load relative to conventional problem-solving, an effect subsequently '
        'termed the "worked example effect." Adaptive systems that dynamically select the '
        'appropriate level and type of scaffolding based on learner performance hold promise '
        'for replicating and extending these benefits at scale.'
    )
    doc.add_paragraph(
        '2.4 Gaps in the Literature\n'
        'Despite the volume of published research on CLT and adaptive learning, several important '
        'gaps remain. First, longitudinal studies tracking cognitive load dynamics across extended '
        'learning periods are scarce. The majority of existing studies are conducted within a '
        'single session, limiting understanding of how load profiles evolve as learners develop '
        'expertise. Second, cross-modal validation of cognitive load measures within adaptive '
        'learning contexts has received limited attention. Establishing convergent validity '
        'across subjective, behavioral, and physiological indices is essential for advancing '
        'theoretical precision and practical measurement guidance.'
    )

    doc.add_page_break()

    # ===== Chapter 3: Methodology =====
    add_heading_with_bookmark(doc, 'Chapter 3: Methodology', 1, '_Chapter3_Methodology')

    doc.add_paragraph(
        'This chapter describes the research design, participant characteristics, data collection '
        'procedures, and analytical methods employed in the present study. A sequential '
        'explanatory mixed-methods design was adopted, in which quantitative data were collected '
        'and analyzed first, followed by qualitative data that helped explain and elaborate '
        'quantitative findings.'
    )
    doc.add_paragraph(
        '3.1 Participants\n'
        'Participants were 142 undergraduate students (78 female, 64 male; mean age = 20.3 years, '
        'SD = 1.8) recruited from introductory psychology and educational technology courses at '
        'two mid-sized research universities in the northeastern United States. Inclusion criteria '
        'required normal or corrected-to-normal vision, no history of neurological disorder, '
        'and baseline computer literacy sufficient for unsupported online learning. Participants '
        'provided written informed consent in accordance with institutional review board approval '
        '(Protocol IRB-2022-0451).'
    )
    doc.add_paragraph(
        '3.2 Learning Platform and Content\n'
        'All participants engaged with a custom-built adaptive learning platform developed in '
        'collaboration with the University Center for Educational Technology. The platform '
        'presented introductory content in two domains: statistical reasoning and argumentation '
        'theory. Adaptive scaffolding was triggered by real-time performance monitoring: when '
        'a learner\'s response latency or accuracy fell below adaptive thresholds, the system '
        'delivered context-sensitive hints or modulated task difficulty.'
    )
    doc.add_paragraph(
        '3.3 Physiological Data Collection\n'
        'Cognitive load was indexed through three parallel measurement channels. Eye-tracking '
        'data were recorded using a Tobii Pro Fusion eye tracker (120 Hz sampling rate) '
        'mounted below the display. EEG was recorded from a 32-channel active electrode cap '
        '(BrainProducts actiCHamp) with a sampling rate of 500 Hz. Subjective load ratings '
        'were collected via an adapted NASA Task Load Index (NASA-TLX) administered between '
        'learning segments.'
    )
    doc.add_paragraph(
        '3.4 Procedure\n'
        'Each experimental session spanned approximately 90 minutes across three visits scheduled '
        'one week apart. During the first session, participants completed baseline assessments '
        'of prior knowledge, working memory capacity (Operation Span task; Unsworth et al., '
        '2005), and trait anxiety (STAI Form Y-2; Spielberger et al., 1983). Subsequent '
        'sessions involved 45-minute adaptive learning modules followed by immediate recall '
        'tests and delayed retention assessments administered 72 hours after each session.'
    )

    # Page 3 content continues with the critical paragraph
    doc.add_paragraph(
        '3.5 Data Analysis\n'
        'Quantitative data were analyzed using a combination of multilevel modeling and '
        'repeated-measures ANOVA to account for the nested structure of observations within '
        'individuals. EEG preprocessing followed established guidelines including bandpass '
        'filtering (1-40 Hz), independent component analysis for artifact removal, and '
        'time-frequency decomposition using the Morlet wavelet transform. Qualitative '
        'interview data were analyzed thematically following Braun and Clarke (2006).'
    )
    doc.add_paragraph(
        '3.6 Limitations\n'
        'Several methodological limitations warrant acknowledgment. First, the sample was '
        'drawn from two universities with similar demographic profiles, potentially limiting '
        'generalizability. Second, EEG sensitivity to motion artifacts introduced data loss '
        'for approximately 12% of participants during select recording windows. Third, the '
        'adaptive platform was purpose-built for this study and may not generalize to '
        'commercial ALS implementations. These considerations are addressed more fully '
        'in the concluding analysis.'
    )

    # THE CRITICAL PARAGRAPH - ends with an open parenthesis for the cross-reference
    critical_para = doc.add_paragraph()
    critical_para.add_run(
        'Notwithstanding these limitations, the methodology provides a rigorous foundation '
        'for testing our core hypotheses. The multi-modal measurement approach captures '
        'cognitive load with greater fidelity than single-measure designs, and the '
        'longitudinal structure allows tracking of load adaptation over time. '
        'The implications are explored later ('
    )
    # NOTE: The cross-reference will be inserted here in golden_patch.py
    # The paragraph intentionally ends with an open parenthesis, awaiting the cross-reference

    doc.add_page_break()

    # ===== Chapter 4: Discussion =====
    add_heading_with_bookmark(doc, 'Chapter 4: Discussion', 1, '_Chapter4_Discussion')

    doc.add_paragraph(
        'The findings of this study carry substantial theoretical and practical implications '
        'for the field of educational technology and instructional design. This chapter '
        'examines these implications in depth, situating our results within existing '
        'theoretical frameworks and highlighting avenues for future investigation.'
    )
    doc.add_paragraph(
        '4.1 Theoretical Contributions\n'
        'Our data provide robust longitudinal support for the core tenets of Cognitive Load '
        'Theory while extending the framework in several important respects. The differential '
        'trajectory of extraneous versus germane load across the three-semester period '
        'suggests that adaptive scaffolding operates primarily by reducing extraneous '
        'processing burden, allowing cognitive resources to be redirected toward '
        'schema-building activities. This pattern aligns with predictions derived from '
        'the "element interactivity" account of CLT (Sweller, 2010) and suggests '
        'that effective adaptive systems may lower the element interactivity threshold '
        'required for deep learning.'
    )
    doc.add_paragraph(
        '4.2 Practical Implications for Instructional Design\n'
        'The present findings yield several actionable recommendations for designers of '
        'adaptive learning systems. First, adaptive hint systems should be calibrated to '
        'trigger based on combined response latency and accuracy metrics rather than '
        'accuracy alone, as our data indicate that elevated latency often precedes '
        'accuracy decline and may signal incipient cognitive overload. Second, the '
        'modality of scaffolding feedback matters: auditory hints were associated with '
        'lower extraneous load increments than text-based hints for material with high '
        'visual complexity, consistent with the modality effect documented in CLT research '
        '(Moreno & Mayer, 1999).'
    )
    doc.add_paragraph(
        '4.3 Cognitive Load Trajectories Across Learner Cohorts\n'
        'A particularly striking finding was the divergence between high-performing and '
        'low-performing cohorts in their germane load trajectories. High performers showed '
        'a characteristic "load inversion" pattern—initially elevated germane load that '
        'declined as content was mastered and schema automation occurred. Low performers '
        'showed persistently elevated extraneous load and suppressed germane load, '
        'suggesting that adaptive scaffolding was insufficient to redirect their cognitive '
        'resources toward constructive schema formation. These findings echo prior work '
        'by Kalyuga et al. (2003) on expertise reversal effects and suggest that '
        'personalization algorithms may need to account for learner-specific load profiles.'
    )
    doc.add_paragraph(
        '4.4 EEG and Pupillometric Convergent Validity\n'
        'The convergent validity between EEG-derived frontal theta power and pupil '
        'dilation indices was moderately strong (r = 0.61, p < 0.001), providing '
        'cross-modal confirmation that both measures captured meaningful variance in '
        'cognitive load. Discrepancies between the measures were most pronounced during '
        'transitions between task segments, likely reflecting differential sensitivity '
        'to arousal and attentional shifts. Future research should investigate the '
        'temporal dynamics of load-related EEG and pupillometric responses with '
        'higher time resolution.'
    )
    doc.add_paragraph(
        '4.5 Platform Engagement and Load Dynamics\n'
        'Analysis of platform interaction logs revealed that scaffolding engagement rates '
        'were negatively correlated with concurrent extraneous load (r = -0.47), '
        'suggesting that learners experiencing high load were less likely to utilize '
        'available scaffolding resources—a paradox with significant design implications. '
        'Future adaptive systems might benefit from "push" scaffolding delivery rather '
        'than "pull" mechanisms, proactively presenting support when load indicators '
        'surpass threshold rather than requiring voluntary learner action.'
    )

    doc.add_page_break()

    # ===== Chapter 5: Conclusion =====
    add_heading_with_bookmark(doc, 'Chapter 5: Conclusion', 1, '_Chapter5_Conclusion')

    doc.add_paragraph(
        'This dissertation has presented a comprehensive longitudinal investigation of '
        'cognitive load dynamics in adaptive learning environments. Using a multi-modal '
        'measurement approach integrating EEG, eye-tracking, and self-report instruments, '
        'we have demonstrated that adaptive scaffolding reliably reduces extraneous '
        'cognitive load while preserving and, under optimal conditions, enhancing '
        'germane load allocation.'
    )
    doc.add_paragraph(
        '5.1 Summary of Key Findings\n'
        'Across three semesters of longitudinal measurement, we observed consistent '
        'reductions in extraneous cognitive load in the adaptive scaffolding condition '
        'relative to controls (mean reduction: 34%, 95% CI [28%, 41%]). Germane load '
        'was maintained or slightly elevated in high-performing cohorts, particularly '
        'during phases of active schema construction. The load inversion pattern among '
        'high performers represents a novel empirical contribution, extending prior work '
        'on expertise reversal to the context of adaptive digital environments.'
    )
    doc.add_paragraph(
        '5.2 Future Directions\n'
        'Several promising avenues for future research emerge from this investigation. '
        'First, the development of real-time adaptive systems that incorporate '
        'physiological load indices directly into their adaptation algorithms represents '
        'an important engineering and design challenge. Second, cross-cultural replication '
        'of the present findings is necessary to establish generalizability across diverse '
        'learner populations. Third, the long-term retention and transfer effects of '
        'load-optimized instruction remain understudied and warrant systematic examination '
        'in future longitudinal designs.'
    )
    doc.add_paragraph(
        '5.3 Closing Remarks\n'
        'The convergence of adaptive learning technology with cognitive neuroscience '
        'methods opens transformative possibilities for educational research and practice. '
        'As measurement technologies become increasingly unobtrusive and computationally '
        'tractable, the vision of truly individualized instruction calibrated to real-time '
        'cognitive load dynamics moves closer to realization. We hope this work contributes '
        'meaningfully to that ongoing project, and that the methodological framework '
        'developed here serves as a useful template for subsequent investigations.'
    )

    doc.add_page_break()

    # ===== References =====
    doc.add_heading('References', level=1)

    references = [
        'Anderson, J. R., Corbett, A. T., Koedinger, K. R., & Pelletier, R. (1995). Cognitive tutors: Lessons learned. Journal of the Learning Sciences, 4(2), 167-207.',
        'Beatty, J. (1982). Task-evoked pupillary responses, processing load, and the structure of processing resources. Psychological Bulletin, 91(2), 276-292.',
        'Braun, V., & Clarke, V. (2006). Using thematic analysis in psychology. Qualitative Research in Psychology, 3(2), 77-101.',
        'Carbonell, J. R. (1970). AI in CAI: An artificial intelligence approach to computer-aided instruction. IEEE Transactions on Man-Machine Systems, 11(4), 190-202.',
        'Kahneman, D., & Beatty, J. (1966). Pupil diameter and load on memory. Science, 154(3756), 1583-1585.',
        'Kalyuga, S., Ayres, P., Chandler, P., & Sweller, J. (2003). The expertise reversal effect. Educational Psychologist, 38(1), 23-31.',
        'Klimesch, W. (1999). EEG alpha and theta oscillations reflect cognitive and memory performance: A review and analysis. Brain Research Reviews, 29(2-3), 169-195.',
        'Moreno, R., & Mayer, R. E. (1999). Cognitive principles of multimedia learning: The role of modality and contiguity. Journal of Educational Psychology, 91(2), 358-368.',
        'Paas, F. G. (1992). Training strategies for attaining transfer of problem-solving skill in statistics: A cognitive-load approach. Journal of Educational Psychology, 84(4), 429-434.',
        'Paas, F., & van Merriënboer, J. J. G. (1994). Variability of worked examples and transfer of geometrical problem-solving skills. Journal of Educational Psychology, 86(1), 122-133.',
        'Spielberger, C. D., Gorsuch, R. L., Lushene, R., Vagg, P. R., & Jacobs, G. A. (1983). Manual for the State-Trait Anxiety Inventory. Consulting Psychologists Press.',
        'Sweller, J. (1988). Cognitive load during problem solving: Effects on learning. Cognitive Science, 12(2), 257-285.',
        'Sweller, J. (2010). Element interactivity and intrinsic, extraneous, and germane cognitive load. Educational Psychology Review, 22(2), 123-138.',
        'Sweller, J., & Cooper, G. A. (1985). The use of worked examples as a substitute for problem solving in learning algebra. Cognition and Instruction, 2(1), 59-89.',
        'Unsworth, N., Heitz, R. P., Schrock, J. C., & Engle, R. W. (2005). An automated version of the operation span task. Behavior Research Methods, 37(3), 498-505.',
        'Wood, D., Bruner, J. S., & Ross, G. (1976). The role of tutoring in problem solving. Journal of Child Psychology and Psychiatry, 17(2), 89-100.',
    ]

    for ref in references:
        p = doc.add_paragraph(ref, style='Normal')
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
