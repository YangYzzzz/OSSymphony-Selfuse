"""
Initial Setup: Add subject and keywords to document properties
Task ID: writer_struct_016
Domain: libreoffice_writer

Creates a 5-page financial report document at /home/user/Desktop/financial_review.docx
with document properties:
  - title = 'Financial Review'
  - author = 'CFO Office'
  - subject = '' (empty — agent must fill this in)
  - keywords = '' (empty — agent must fill this in)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'financial_review'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_doc_property(doc, prop_name, value):
    """Set a core document property via XML manipulation."""
    core_props = doc.core_properties
    setattr(core_props, prop_name, value)


def create_initial():
    # Ensure the Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set document properties
    doc.core_properties.title = 'Financial Review'
    doc.core_properties.author = 'CFO Office'
    # Subject and keywords intentionally left empty
    doc.core_properties.subject = ''
    doc.core_properties.keywords = ''

    # --- Page 1: Executive Summary ---
    heading = doc.add_heading('Financial Review', level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub_heading = doc.add_paragraph('Q4 2025 Executive Summary')
    sub_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in sub_heading.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph('')

    doc.add_heading('Executive Summary', level=1)

    p = doc.add_paragraph(
        'This report provides a comprehensive analysis of the company\'s financial performance '
        'for the fourth quarter of fiscal year 2025. Overall revenue reached $2.4 billion, '
        'representing a 12% increase year-over-year. Operating expenses were well-controlled '
        'at $1.7 billion, resulting in an operating margin of 29.2%.'
    )

    doc.add_paragraph(
        'Net income for Q4 2025 totaled $486 million, up from $412 million in Q4 2024. '
        'Earnings per share (EPS) improved to $3.24, exceeding analyst consensus estimates '
        'of $3.10. Cash flow from operations remained robust at $612 million.'
    )

    doc.add_heading('Key Performance Indicators', level=2)

    table1 = doc.add_table(rows=6, cols=3)
    table1.style = 'Table Grid'
    headers = ['Metric', 'Q4 2025', 'Q4 2024']
    for i, h in enumerate(headers):
        cell = table1.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    kpi_data = [
        ('Revenue', '$2,412M', '$2,152M'),
        ('Operating Income', '$704M', '$618M'),
        ('Net Income', '$486M', '$412M'),
        ('EPS (Diluted)', '$3.24', '$2.75'),
        ('Free Cash Flow', '$498M', '$433M'),
    ]
    for row_idx, (metric, q4_2025, q4_2024) in enumerate(kpi_data, start=1):
        table1.cell(row_idx, 0).text = metric
        table1.cell(row_idx, 1).text = q4_2025
        table1.cell(row_idx, 2).text = q4_2024

    doc.add_page_break()

    # --- Page 2: Revenue Analysis ---
    doc.add_heading('Revenue Analysis', level=1)

    doc.add_paragraph(
        'Total revenue for Q4 2025 was $2,412 million, driven by strong performance across '
        'all business segments. The Americas region contributed the largest share at 48%, '
        'followed by EMEA at 32% and APAC at 20%.'
    )

    doc.add_heading('Revenue by Segment', level=2)
    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'
    seg_headers = ['Segment', 'Revenue ($M)', 'Growth YoY', 'Margin']
    for i, h in enumerate(seg_headers):
        run = table2.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True

    seg_data = [
        ('Enterprise Solutions', '$1,056M', '+15.2%', '34.1%'),
        ('Consumer Products', '$724M', '+8.7%', '22.5%'),
        ('Cloud Services', '$438M', '+28.4%', '41.6%'),
        ('Professional Services', '$194M', '+5.1%', '18.9%'),
    ]
    for row_idx, row in enumerate(seg_data, start=1):
        for col_idx, val in enumerate(row):
            table2.cell(row_idx, col_idx).text = val

    doc.add_paragraph('')
    doc.add_paragraph(
        'Cloud Services continued to be the fastest-growing segment with 28.4% YoY growth, '
        'reflecting the company\'s strategic investment in cloud infrastructure and digital '
        'transformation services. Enterprise Solutions remained the largest revenue contributor.'
    )

    doc.add_heading('Geographic Performance', level=2)
    doc.add_paragraph(
        'The Americas region generated $1,158M in revenue, up 11.3% from the prior year. '
        'EMEA delivered $771M, growing 13.8% driven by strong enterprise demand in Western Europe. '
        'APAC revenue reached $483M, representing 14.2% growth, with particularly strong results '
        'in Japan and Australia.'
    )

    doc.add_page_break()

    # --- Page 3: Operating Expenses ---
    doc.add_heading('Operating Expenses', level=1)

    doc.add_paragraph(
        'Total operating expenses for Q4 2025 were $1,708 million, representing 70.8% of revenue. '
        'This compares favorably to 71.3% in Q4 2024, reflecting ongoing efficiency improvements '
        'and economies of scale.'
    )

    doc.add_heading('Expense Breakdown', level=2)
    table3 = doc.add_table(rows=6, cols=3)
    table3.style = 'Table Grid'
    exp_headers = ['Category', 'Q4 2025 ($M)', '% of Revenue']
    for i, h in enumerate(exp_headers):
        run = table3.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True

    exp_data = [
        ('Cost of Goods Sold', '$892M', '37.0%'),
        ('Research & Development', '$312M', '12.9%'),
        ('Sales & Marketing', '$286M', '11.9%'),
        ('General & Administrative', '$218M', '9.0%'),
    ]
    for row_idx, row in enumerate(exp_data, start=1):
        for col_idx, val in enumerate(row):
            table3.cell(row_idx, col_idx).text = val

    doc.add_paragraph('')
    doc.add_paragraph(
        'Research and development spending increased by $28M compared to Q4 2024, reflecting '
        'the company\'s continued commitment to innovation and product development. R&D as a '
        'percentage of revenue declined slightly from 13.2% to 12.9%, demonstrating improved efficiency.'
    )

    doc.add_heading('Headcount & Workforce', level=2)
    doc.add_paragraph(
        'Total global headcount at the end of Q4 2025 was 18,432 employees, an increase of 1,243 '
        'from the same period last year. The majority of new hires were in engineering and cloud '
        'services roles, supporting the company\'s growth strategy. Employee retention rate '
        'improved to 91.4% from 89.8% in 2024.'
    )

    doc.add_page_break()

    # --- Page 4: Balance Sheet & Cash Flow ---
    doc.add_heading('Balance Sheet & Cash Flow', level=1)

    doc.add_paragraph(
        'The company maintained a strong balance sheet with total assets of $12.6 billion as of '
        'December 31, 2025. Cash and cash equivalents stood at $2.1 billion, providing ample '
        'liquidity for strategic initiatives and shareholder returns.'
    )

    doc.add_heading('Cash Flow Summary', level=2)
    table4 = doc.add_table(rows=5, cols=3)
    table4.style = 'Table Grid'
    cf_headers = ['Cash Flow Component', 'Q4 2025 ($M)', 'Q4 2024 ($M)']
    for i, h in enumerate(cf_headers):
        run = table4.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True

    cf_data = [
        ('Operating Cash Flow', '$612M', '$541M'),
        ('Capital Expenditures', '($114M)', '($108M)'),
        ('Free Cash Flow', '$498M', '$433M'),
        ('Share Repurchases', '($245M)', '($198M)'),
    ]
    for row_idx, row in enumerate(cf_data, start=1):
        for col_idx, val in enumerate(row):
            table4.cell(row_idx, col_idx).text = val

    doc.add_paragraph('')
    doc.add_paragraph(
        'Free cash flow of $498M was used primarily for share repurchases ($245M), dividends '
        '($112M), and strategic acquisitions ($87M). The remaining $54M was retained to strengthen '
        'the balance sheet.'
    )

    doc.add_heading('Capital Allocation', level=2)
    doc.add_paragraph(
        'The Board of Directors approved a quarterly dividend of $0.75 per share, representing '
        'a 7.1% increase over the prior year dividend of $0.70 per share. The company\'s share '
        'repurchase program remains active with $1.2 billion remaining under the current authorization.'
    )

    doc.add_page_break()

    # --- Page 5: Outlook & Conclusion ---
    doc.add_heading('2026 Outlook & Conclusion', level=1)

    doc.add_heading('Financial Guidance', level=2)
    doc.add_paragraph(
        'For fiscal year 2026, the company provides the following guidance:'
    )

    table5 = doc.add_table(rows=5, cols=3)
    table5.style = 'Table Grid'
    guid_headers = ['Metric', 'FY2026 Guidance', 'FY2025 Actual']
    for i, h in enumerate(guid_headers):
        run = table5.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True

    guid_data = [
        ('Revenue', '$9.8B - $10.2B', '$9.3B'),
        ('Operating Margin', '29% - 31%', '29.2%'),
        ('EPS (Diluted)', '$12.80 - $13.40', '$12.18'),
        ('Capital Expenditures', '$440M - $480M', '$458M'),
    ]
    for row_idx, row in enumerate(guid_data, start=1):
        for col_idx, val in enumerate(row):
            table5.cell(row_idx, col_idx).text = val

    doc.add_paragraph('')

    doc.add_heading('Strategic Priorities', level=2)
    doc.add_paragraph(
        'Management has identified four strategic priorities for 2026:'
    )
    doc.add_paragraph('Accelerate cloud services growth through expanded partner ecosystem', style='List Number')
    doc.add_paragraph('Invest in AI-powered product capabilities across all segments', style='List Number')
    doc.add_paragraph('Expand geographic presence in APAC emerging markets', style='List Number')
    doc.add_paragraph('Continue operational efficiency programs targeting $150M in cost savings', style='List Number')

    doc.add_heading('Conclusion', level=2)
    doc.add_paragraph(
        'Q4 2025 demonstrated the company\'s ability to deliver strong financial performance '
        'while investing for long-term growth. Revenue growth of 12%, margin expansion, and '
        'robust free cash flow generation reflect the strength of our business model and '
        'disciplined execution.'
    )
    doc.add_paragraph(
        'We enter 2026 with confidence, supported by a strong balance sheet, talented workforce, '
        'and clear strategic direction. The CFO Office thanks shareholders, employees, and '
        'business partners for their continued support.'
    )

    p_footer = doc.add_paragraph('')
    p_footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p_footer.add_run('Prepared by: CFO Office | Confidential — For Internal Use Only')
    run.font.size = Pt(9)
    run.font.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
