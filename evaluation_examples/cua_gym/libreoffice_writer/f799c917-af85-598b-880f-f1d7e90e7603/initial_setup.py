"""
Initial Setup: Set up an AutoCorrect entry that replaces 'sig1' with a formatted signature
Task ID: writer_frd_052
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_052'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Liberation Serif'
    font.size = Pt(12)

    # Add a title paragraph
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(18)
    title_run = title_para.add_run('Quarterly Research Report')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = 'Liberation Sans'

    # Add some introductory text to make the document realistic
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    run = intro.add_run(
        'This report summarizes the key findings and progress from Q1 2026 across '
        'all active research projects. Each section below provides an overview of '
        'milestones achieved, challenges encountered, and planned next steps.'
    )
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)

    # Add section heading
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    h_run = heading.add_run('Project Alpha - Neural Architecture Search')
    h_run.bold = True
    h_run.font.size = Pt(14)
    h_run.font.name = 'Liberation Sans'

    # Add project details
    details = [
        'The team successfully completed Phase 2 of the neural architecture search pipeline, '
        'achieving a 15% improvement in model efficiency compared to baseline measurements.',
        'Key deliverables included the automated hyperparameter tuning framework and the '
        'distributed training infrastructure upgrade. Both components passed integration testing.',
        'Budget utilization stands at 78% with two months remaining in the fiscal quarter. '
        'We anticipate staying within the allocated $2.4M for this project phase.',
    ]

    for text in details:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = 'Liberation Serif'
        r.font.size = Pt(12)

    # Add another section
    heading2 = doc.add_paragraph()
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)
    h2_run = heading2.add_run('Project Beta - Robotic Process Automation')
    h2_run.bold = True
    h2_run.font.size = Pt(14)
    h2_run.font.name = 'Liberation Sans'

    details2 = [
        'The RPA initiative has onboarded 12 new processes for automation during Q1, '
        'bringing the total automated workflows to 47 across three departments.',
        'Estimated annual cost savings from newly automated processes total approximately '
        '$890,000, with an average error rate reduction of 94% compared to manual handling.',
    ]

    for text in details2:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = 'Liberation Serif'
        r.font.size = Pt(12)

    # Add a blank line before the signature block
    doc.add_paragraph()

    # Add the formatted signature block that the user will use
    # This is the text that should be selected and used as the AutoCorrect entry
    sig_line1 = doc.add_paragraph()
    sig_line1.paragraph_format.space_after = Pt(0)
    name_run = sig_line1.add_run('Dr. Robert Chen')
    name_run.bold = True
    name_run.font.name = 'Liberation Serif'
    name_run.font.size = Pt(12)

    sig_line2 = doc.add_paragraph()
    sig_line2.paragraph_format.space_before = Pt(0)
    title_run2 = sig_line2.add_run('Director of Research')
    title_run2.italic = True
    title_run2.font.name = 'Liberation Serif'
    title_run2.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
