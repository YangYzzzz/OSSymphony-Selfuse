"""
Initial Setup: Engagement letter template document (no macros)
Task ID: writer_legal_063
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_063'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Firm Letterhead ---
    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.add_run("MORRISON & WHITFIELD LLP")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    sub_heading = doc.add_paragraph()
    sub_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sub_heading.add_run("Attorneys at Law")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.italic = True

    address = doc.add_paragraph()
    address.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = address.add_run("350 Park Avenue, Suite 2400\nNew York, NY 10022\nTel: (212) 555-8700 | Fax: (212) 555-8701")
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"

    # Horizontal line (border on paragraph)
    hr_para = doc.add_paragraph()
    hr_para.paragraph_format.space_before = Pt(6)
    hr_para.paragraph_format.space_after = Pt(12)
    pPr = hr_para._element.get_or_add_pPr()
    from docx.oxml.ns import qn
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '000000',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    # --- Date and Addressee ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(12)
    run = date_para.add_run("March 28, 2026")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    # Delivery method
    delivery = doc.add_paragraph()
    run = delivery.add_run("VIA EMAIL AND FIRST CLASS MAIL")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    # Addressee block
    addr_lines = [
        "Ms. Catherine A. Reeves",
        "Chief Executive Officer",
        "Pinnacle Healthcare Systems, Inc.",
        "1200 Corporate Boulevard, Suite 500",
        "White Plains, NY 10601",
    ]
    for line in addr_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    # Spacer
    doc.add_paragraph()

    # --- Re line ---
    re_para = doc.add_paragraph()
    run = re_para.add_run("Re: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run = re_para.add_run("Engagement Letter — Regulatory Compliance Review and Advisory Services")
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    # --- Salutation ---
    sal = doc.add_paragraph()
    sal.paragraph_format.space_before = Pt(12)
    run = sal.add_run("Dear Ms. Reeves:")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    # --- Body paragraphs ---
    body_texts = [
        "Thank you for selecting Morrison & Whitfield LLP to represent Pinnacle Healthcare Systems, Inc. (\"Pinnacle\" or the \"Company\") in connection with the regulatory compliance review and advisory services described below. This letter sets forth the terms under which we will provide legal services to the Company.",

        "1.  Scope of Engagement. We have been retained to provide legal counsel and advisory services related to the Company's compliance with applicable federal and state healthcare regulations, including but not limited to the Health Insurance Portability and Accountability Act (\"HIPAA\"), the Anti-Kickback Statute (42 U.S.C. § 1320a-7b), and the Stark Law (42 U.S.C. § 1395nn). Our services will include a comprehensive review of the Company's existing compliance programs, identification of potential areas of regulatory risk, and preparation of remedial recommendations.",

        "2.  Responsible Attorneys. The primary attorneys responsible for this engagement will be David R. Morrison, Partner, and Jennifer L. Whitfield, Of Counsel. Other attorneys and legal professionals within the firm may assist with the engagement as needed. You will be informed in advance of any material change in the staffing of your matter.",

        "3.  Fees and Billing. Our fees for this engagement will be based on our standard hourly rates, which currently range from $425 to $875 per hour depending on the seniority and experience of the attorney or professional involved. We will provide monthly invoices detailing the services performed, the time spent, and any disbursements incurred. Payment is due within thirty (30) days of the invoice date.",

        "4.  Retainer. Upon execution of this engagement letter, we request an initial retainer in the amount of Fifty Thousand Dollars ($50,000.00). The retainer will be held in our client trust account and applied against fees and expenses as they are incurred. We may request replenishment of the retainer from time to time as the matter progresses.",

        "5.  Expenses. In addition to professional fees, the Company will be responsible for reimbursement of reasonable out-of-pocket expenses incurred in connection with this engagement, including but not limited to filing fees, court costs, travel expenses, document reproduction, expert witness fees, and electronic discovery vendors.",

        "6.  Confidentiality. All communications between Morrison & Whitfield LLP and Pinnacle Healthcare Systems, Inc. in connection with this engagement are protected by the attorney-client privilege and work product doctrine. We will maintain the confidentiality of all information provided to us in the course of our representation.",

        "7.  Conflicts of Interest. We have conducted a conflicts check and are not aware of any conflicts that would preclude us from representing the Company in this matter. Should a conflict arise during the course of our engagement, we will promptly notify you and take appropriate steps in accordance with our professional obligations.",

        "8.  Termination. Either party may terminate this engagement at any time upon written notice to the other party. In the event of termination, the Company will remain responsible for payment of all fees and expenses incurred through the date of termination. Upon termination, we will promptly return all original documents and property belonging to the Company.",
    ]

    for text in body_texts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    # Note: No closing paragraph here — that's what the macro task requires

    # --- Signature block ---
    sig_spacer = doc.add_paragraph()
    sig_spacer.paragraph_format.space_before = Pt(24)
    run = sig_spacer.add_run("Very truly yours,")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    # Signature space
    for _ in range(3):
        doc.add_paragraph()

    sig_name = doc.add_paragraph()
    run = sig_name.add_run("David R. Morrison")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_after = Pt(0)
    run = sig_title.add_run("Managing Partner")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    sig_firm = doc.add_paragraph()
    run = sig_firm.add_run("Morrison & Whitfield LLP")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    # --- Acknowledgment block ---
    doc.add_paragraph()
    ack_header = doc.add_paragraph()
    ack_header.paragraph_format.space_before = Pt(12)
    run = ack_header.add_run("AGREED AND ACCEPTED:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    ack_line = doc.add_paragraph()
    run = ack_line.add_run("_________________________________")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    ack_name = doc.add_paragraph()
    ack_name.paragraph_format.space_after = Pt(0)
    run = ack_name.add_run("Catherine A. Reeves")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    ack_title = doc.add_paragraph()
    ack_title.paragraph_format.space_after = Pt(0)
    run = ack_title.add_run("Chief Executive Officer")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    ack_company = doc.add_paragraph()
    run = ack_company.add_run("Pinnacle Healthcare Systems, Inc.")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    ack_date = doc.add_paragraph()
    run = ack_date.add_run("Date: _______________")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
