"""
Initial Setup: Document with a text box that has a thin solid black border
Task ID: writer_obj_022
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_022'
# Task context says file is at ~/Desktop/sidebar_doc.docx
OUTPUT = f'{WORKDIR}/Desktop/sidebar_doc.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def cm_to_emu(cm_val):
    """Convert centimeters to EMU (English Metric Units)."""
    return int(cm_val * 914400 / 2.54)


def create_textbox_with_border(doc):
    """
    Create a text box (5cm x 8cm) on the right side of page 1 with a thin solid black border.
    Text boxes in .docx are created via drawing shapes in XML.
    """
    # We insert a floating text box using wps:wsp (Word Processing Shape) in a paragraph
    # The text box is positioned to the right side of the page

    # Width = 5cm, Height = 8cm in EMU
    width_emu = cm_to_emu(5)   # 5cm
    height_emu = cm_to_emu(8)  # 8cm

    # Position: right side, near top of page
    # posOffset from right edge: ~1cm from right margin
    pos_x_emu = cm_to_emu(12)  # ~12cm from left (approx right side for A4)
    pos_y_emu = cm_to_emu(3)   # 3cm from top

    # Build the XML for the drawing element with text box
    drawing_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:r>
    <w:rPr/>
    <w:drawing>
      <wp:anchor distT="114300" distB="114300" distL="114300" distR="114300"
                 simplePos="0" relativeHeight="251658240" behindDoc="0"
                 locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column">
          <wp:posOffset>{pos_x_emu}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="paragraph">
          <wp:posOffset>{pos_y_emu}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapTopAndBottom/>
        <wp:docPr id="1" name="Text Box 1"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp>
              <wps:cNvSpPr txBx="1">
                <a:spLocks noChangeArrowheads="1"/>
              </wps:cNvSpPr>
              <wps:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{width_emu}" cy="{height_emu}"/>
                </a:xfrm>
                <a:prstGeom prst="rect">
                  <a:avLst/>
                </a:prstGeom>
                <a:solidFill>
                  <a:srgbClr val="FFFFFF"/>
                </a:solidFill>
                <a:ln w="12700">
                  <a:solidFill>
                    <a:srgbClr val="000000"/>
                  </a:solidFill>
                </a:ln>
              </wps:spPr>
              <wps:txbx>
                <w:txbxContent>
                  <w:p>
                    <w:pPr>
                      <w:jc w:val="left"/>
                    </w:pPr>
                    <w:r>
                      <w:rPr>
                        <w:b/>
                        <w:sz w:val="24"/>
                      </w:rPr>
                      <w:t>Project Overview</w:t>
                    </w:r>
                  </w:p>
                  <w:p>
                    <w:r>
                      <w:t xml:space="preserve">Q1 2025 targets have been set. Review the attached performance metrics and align your team goals accordingly.</w:t>
                    </w:r>
                  </w:p>
                  <w:p>
                    <w:r>
                      <w:t xml:space="preserve">Key milestones: March 31 deadline for Phase 1 deliverables.</w:t>
                    </w:r>
                  </w:p>
                </w:txbxContent>
              </wps:txbx>
              <wps:bodyPr insFit="norm">
                <a:noAutofit/>
              </wps:bodyPr>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </w:r>
</w:p>'''

    return drawing_xml


def create_initial():
    doc = Document()

    # Set page to A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Add a title heading
    title = doc.add_heading('Annual Business Report 2025', level=1)

    # Add text box paragraph with border using XML injection
    drawing_xml = create_textbox_with_border(doc)
    drawing_element = etree.fromstring(drawing_xml)
    doc.element.body.append(drawing_element)

    # Add body paragraphs with realistic content
    doc.add_paragraph(
        'This report summarizes the key business outcomes and strategic initiatives '
        'undertaken during the fiscal year 2025. The findings presented herein reflect '
        'extensive analysis across all divisions of the organization.'
    )

    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'Revenue for FY2025 reached $12.4 million, representing a 9.3% increase over '
        'the previous year. Operating margins improved to 18.2%, driven by cost optimization '
        'and new market expansion. Customer retention remained strong at 94.7%.'
    )

    doc.add_heading('Key Performance Indicators', level=2)

    # Add a table with some KPIs
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Metric', 'Target', 'Achieved']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    data_rows = [
        ('Revenue Growth', '8%', '9.3%'),
        ('Operating Margin', '16%', '18.2%'),
        ('Customer Retention', '93%', '94.7%'),
        ('Employee Satisfaction', '80%', '85.1%'),
    ]
    for i, (metric, target, achieved) in enumerate(data_rows, 1):
        table.cell(i, 0).text = metric
        table.cell(i, 1).text = target
        table.cell(i, 2).text = achieved

    doc.add_paragraph('')  # spacing

    doc.add_heading('Strategic Initiatives', level=2)
    doc.add_paragraph(
        'Three major initiatives drove performance improvements this year: the digital '
        'transformation program, the talent development framework, and the customer '
        'experience redesign project.'
    )

    doc.add_paragraph(
        'The digital transformation program invested $2.1 million in new infrastructure, '
        'resulting in a 23% reduction in processing time and improved system reliability '
        'from 99.1% to 99.8% uptime.'
    )

    # Ensure Desktop directory exists
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
