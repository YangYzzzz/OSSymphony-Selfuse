"""
Initial Setup: Brand guidelines document with Pinnacle Solutions, SmartFlow, DataBridge brand names
Task ID: writer_txtfmt_064
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_064'
OUTPUT = f'{WORKDIR}/Desktop/brand_guidelines.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Remove default empty paragraph if present
    # and configure the document with 6 paragraphs

    # Paragraph 1: mentions Pinnacle Solutions
    p1 = doc.add_paragraph()
    r1a = p1.add_run("At ")
    r1a.font.name = "Calibri"
    r1a.font.size = Pt(11)
    r1b = p1.add_run("Pinnacle Solutions")
    r1b.font.name = "Calibri"
    r1b.font.size = Pt(11)
    r1c = p1.add_run(
        ", our brand identity reflects our commitment to excellence and innovation. "
        "We deliver transformative enterprise solutions that empower organizations to "
        "achieve their strategic objectives with confidence and precision."
    )
    r1c.font.name = "Calibri"
    r1c.font.size = Pt(11)

    # Paragraph 2: mentions SmartFlow
    p2 = doc.add_paragraph()
    r2a = p2.add_run(
        "Our flagship product line, "
    )
    r2a.font.name = "Calibri"
    r2a.font.size = Pt(11)
    r2b = p2.add_run("SmartFlow")
    r2b.font.name = "Calibri"
    r2b.font.size = Pt(11)
    r2c = p2.add_run(
        ", represents the next generation of intelligent workflow automation. "
        "The platform seamlessly integrates with existing infrastructure to optimize "
        "business processes and reduce operational overhead across all departments."
    )
    r2c.font.name = "Calibri"
    r2c.font.size = Pt(11)

    # Paragraph 3: mentions DataBridge
    p3 = doc.add_paragraph()
    r3a = p3.add_run(
        "The "
    )
    r3a.font.name = "Calibri"
    r3a.font.size = Pt(11)
    r3b = p3.add_run("DataBridge")
    r3b.font.name = "Calibri"
    r3b.font.size = Pt(11)
    r3c = p3.add_run(
        " suite enables seamless data migration and synchronization between "
        "disparate enterprise systems. Organizations can leverage real-time analytics "
        "to drive informed decision-making and maintain competitive advantage in rapidly evolving markets."
    )
    r3c.font.name = "Calibri"
    r3c.font.size = Pt(11)

    # Paragraph 4: mentions Pinnacle Solutions again
    p4 = doc.add_paragraph()
    r4a = p4.add_run(
        "The visual identity guidelines established by "
    )
    r4a.font.name = "Calibri"
    r4a.font.size = Pt(11)
    r4b = p4.add_run("Pinnacle Solutions")
    r4b.font.name = "Calibri"
    r4b.font.size = Pt(11)
    r4c = p4.add_run(
        " ensure consistency across all marketing materials, digital assets, and "
        "client communications. Brand representatives must adhere strictly to the "
        "approved color palette, typography standards, and logo usage specifications."
    )
    r4c.font.name = "Calibri"
    r4c.font.size = Pt(11)

    # Paragraph 5: mentions SmartFlow again
    p5 = doc.add_paragraph()
    r5a = p5.add_run(
        "Deployment and configuration of "
    )
    r5a.font.name = "Calibri"
    r5a.font.size = Pt(11)
    r5b = p5.add_run("SmartFlow")
    r5b.font.name = "Calibri"
    r5b.font.size = Pt(11)
    r5c = p5.add_run(
        " must follow the certified implementation methodology outlined in the partner "
        "documentation. All customizations require approval from the product architecture "
        "team to maintain platform integrity and ensure long-term supportability."
    )
    r5c.font.name = "Calibri"
    r5c.font.size = Pt(11)

    # Paragraph 6: mentions DataBridge again
    p6 = doc.add_paragraph()
    r6a = p6.add_run(
        "Integration projects utilizing "
    )
    r6a.font.name = "Calibri"
    r6a.font.size = Pt(11)
    r6b = p6.add_run("DataBridge")
    r6b.font.name = "Calibri"
    r6b.font.size = Pt(11)
    r6c = p6.add_run(
        " connectors must undergo rigorous security assessment and compliance validation "
        "prior to production deployment. All data transmission pathways are encrypted "
        "using industry-standard protocols to safeguard sensitive organizational information."
    )
    r6c.font.name = "Calibri"
    r6c.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
