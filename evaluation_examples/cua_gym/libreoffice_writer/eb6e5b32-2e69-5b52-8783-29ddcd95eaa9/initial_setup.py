"""
Initial Setup: Job description document in A4 paper size
Task ID: writer_hr_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set A4 page size (21 x 29.7 cm)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.orientation = WD_ORIENT.PORTRAIT

    # Company header
    heading = doc.add_heading('Meridian Technologies Inc.', level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('JOB DESCRIPTION')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x7A)

    # Separator line
    doc.add_paragraph('_' * 60)

    # Position title
    p = doc.add_paragraph()
    run = p.add_run('Position Title: ')
    run.bold = True
    run.font.size = Pt(12)
    p.add_run('Senior Software Engineer').font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run('Department: ')
    run.bold = True
    p.add_run('Platform Engineering')

    p = doc.add_paragraph()
    run = p.add_run('Reports To: ')
    run.bold = True
    p.add_run('VP of Engineering, David Nakamura')

    p = doc.add_paragraph()
    run = p.add_run('Location: ')
    run.bold = True
    p.add_run('Austin, TX (Hybrid - 3 days on-site)')

    p = doc.add_paragraph()
    run = p.add_run('Employment Type: ')
    run.bold = True
    p.add_run('Full-Time, Exempt')

    p = doc.add_paragraph()
    run = p.add_run('Salary Range: ')
    run.bold = True
    p.add_run('$145,000 - $185,000 annually')

    p = doc.add_paragraph()
    run = p.add_run('Date Posted: ')
    run.bold = True
    p.add_run('March 15, 2025')

    doc.add_paragraph()  # blank line

    # Position Summary
    doc.add_heading('Position Summary', level=1)
    doc.add_paragraph(
        'Meridian Technologies is seeking an experienced Senior Software Engineer to join '
        'our Platform Engineering team. In this role, you will design, develop, and maintain '
        'scalable backend services that power our enterprise resource planning platform used '
        'by over 2,000 mid-market companies. You will collaborate closely with product managers, '
        'UX designers, and fellow engineers to deliver high-quality features on a bi-weekly '
        'release cadence.'
    )

    # Key Responsibilities
    doc.add_heading('Key Responsibilities', level=1)
    responsibilities = [
        'Architect and implement microservices using Go and Python, deployed on Kubernetes clusters across AWS regions',
        'Lead code reviews and mentor junior engineers on best practices for distributed systems design',
        'Design and optimize PostgreSQL and Redis data models to support sub-100ms API response times',
        'Collaborate with the DevOps team to improve CI/CD pipelines, targeting 99.95% deployment success rate',
        'Participate in on-call rotation (one week per month) and conduct post-incident reviews',
        'Write comprehensive unit and integration tests, maintaining minimum 85% code coverage',
        'Contribute to technical documentation and architecture decision records (ADRs)',
        'Evaluate and integrate third-party APIs and services as needed for product features',
    ]
    for item in responsibilities:
        doc.add_paragraph(item, style='List Bullet')

    # Required Qualifications
    doc.add_heading('Required Qualifications', level=1)
    required = [
        "Bachelor's degree in Computer Science, Software Engineering, or related field",
        '5+ years of professional software development experience',
        'Strong proficiency in at least two of: Go, Python, Java, or TypeScript',
        'Experience with containerization (Docker) and orchestration (Kubernetes)',
        'Solid understanding of relational databases and query optimization',
        'Familiarity with event-driven architectures (Kafka, RabbitMQ, or similar)',
        'Excellent problem-solving skills and attention to detail',
        'Strong written and verbal communication skills',
    ]
    for item in required:
        doc.add_paragraph(item, style='List Bullet')

    # Preferred Qualifications
    doc.add_heading('Preferred Qualifications', level=1)
    preferred = [
        "Master's degree in a relevant technical field",
        'Experience with infrastructure-as-code tools (Terraform, Pulumi)',
        'Contributions to open-source projects',
        'Previous experience in a B2B SaaS environment',
        'Knowledge of observability tools (Datadog, Grafana, Prometheus)',
        'Experience with GraphQL API design',
    ]
    for item in preferred:
        doc.add_paragraph(item, style='List Bullet')

    # Benefits
    doc.add_heading('Benefits & Perks', level=1)
    benefits = [
        'Comprehensive health, dental, and vision insurance (100% premium covered for employees)',
        '401(k) with 4% company match, vested immediately',
        '20 days PTO plus 10 company holidays and 5 sick days',
        '$3,000 annual professional development stipend',
        'Home office equipment allowance up to $2,500',
        'Quarterly team offsites and annual company retreat',
        'Parental leave: 16 weeks paid for primary caregivers, 8 weeks for secondary',
    ]
    for item in benefits:
        doc.add_paragraph(item, style='List Bullet')

    # Equal Opportunity
    doc.add_heading('Equal Opportunity Statement', level=1)
    doc.add_paragraph(
        'Meridian Technologies is an equal opportunity employer. We celebrate diversity and are '
        'committed to creating an inclusive environment for all employees. All qualified applicants '
        'will receive consideration for employment without regard to race, color, religion, gender, '
        'gender identity or expression, sexual orientation, national origin, genetics, disability, '
        'age, or veteran status.'
    )

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Meridian Technologies Inc. | 4200 Research Blvd, Suite 300, Austin, TX 78759')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
