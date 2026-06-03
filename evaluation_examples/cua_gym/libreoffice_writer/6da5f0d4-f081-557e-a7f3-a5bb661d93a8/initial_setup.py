"""
Initial Setup: Reading Comprehension Test - Sentence Spacing Task
Task ID: osworld_writer_spacing_007
Domain: libreoffice_writer

Creates a reading comprehension test document with:
- First paragraph: 11-sentence passage as a single block (no sentence separation)
- Subsequent paragraphs: Question prompts to be left untouched
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_spacing_007'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Reading Comprehension Test")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(12)

    # Subtitle / instructions
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run("Passage A: The Amazon Rainforest")
    subtitle_run.bold = True
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(13)
    subtitle_para.paragraph_format.space_after = Pt(8)

    # Instructions for students
    instr_para = doc.add_paragraph()
    instr_run = instr_para.add_run(
        "Directions: Read the following passage carefully. Answer the questions below using evidence from the text."
    )
    instr_run.font.size = Pt(11)
    instr_para.paragraph_format.space_after = Pt(10)

    # First paragraph: 11-sentence passage as a SINGLE block (task requirement: student must split this)
    passage_sentences = [
        "The Amazon rainforest covers over 5.5 million square kilometers.",
        "It is home to more than three million species of plants, animals, and insects.",
        "Scientists discover hundreds of new species in this region each year.",
        "The forest produces approximately 20 percent of the world's oxygen supply.",
        "Many indigenous communities have lived in the Amazon for thousands of years.",
        "Their traditional knowledge of medicinal plants is invaluable to modern medicine.",
        "Deforestation threatens the survival of both wildlife and human communities in the region.",
        "Each year, millions of acres of rainforest are lost to logging and agriculture.",
        "The loss of tree cover contributes significantly to global climate change.",
        "International conservation efforts have slowed but not stopped this destruction.",
        "Protecting the Amazon requires cooperation between governments, communities, and scientists worldwide.",
    ]

    # All 11 sentences in ONE paragraph block — agent must split them
    passage_text = " ".join(passage_sentences)
    passage_para = doc.add_paragraph()
    passage_run = passage_para.add_run(passage_text)
    passage_run.font.size = Pt(11)
    passage_para.paragraph_format.space_after = Pt(12)
    passage_para.paragraph_format.left_indent = Inches(0)

    # Blank line between passage and questions
    doc.add_paragraph()

    # Question prompts — these should remain untouched
    questions_header = doc.add_paragraph()
    q_header_run = questions_header.add_run("Questions:")
    q_header_run.bold = True
    q_header_run.font.size = Pt(12)
    questions_header.paragraph_format.space_after = Pt(6)

    questions = [
        "1. What is the main idea of the passage? Use at least two details from the text to support your answer.",
        "2. According to the passage, why is the traditional knowledge of indigenous communities important? Explain in your own words.",
        "3. What are two causes of deforestation mentioned in the passage? How do they affect the environment?",
        "4. The passage states that the Amazon produces 20 percent of the world's oxygen. Why is this fact significant?",
        "5. What does the author suggest is necessary to protect the Amazon rainforest? Do you agree? Why or why not?",
    ]

    for q in questions:
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(q)
        q_run.font.size = Pt(11)
        q_para.paragraph_format.space_after = Pt(18)  # space for student answer

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
