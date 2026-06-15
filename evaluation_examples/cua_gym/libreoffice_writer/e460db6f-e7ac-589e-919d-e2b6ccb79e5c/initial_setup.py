"""
Initial Setup: Insert separator pages between subdocuments in a master document.
Task ID: writer_rm_090
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_090'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard letter
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Document title page
    title = doc.add_heading('The Echoes of Tomorrow', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('A Novel by Elena Marchetti')
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Add blank space after title
    for _ in range(3):
        doc.add_paragraph('')

    pub_info = doc.add_paragraph()
    pub_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = pub_info.add_run('Riverstone Publishing House\nFirst Edition, 2025')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Chapter definitions with realistic content
    chapters = [
        {
            'title': 'The Beginning',
            'paragraphs': [
                'The morning sun cast long shadows across the cobblestone streets of Verona as Clara Ashford stepped off the train. She clutched her leather satchel close, feeling the weight of the manuscript inside—the one her grandmother had entrusted to her just days before passing.',
                'The station was nearly empty at this hour. A lone vendor arranged newspapers on a wooden stand, and somewhere in the distance a church bell struck seven. Clara inhaled deeply, tasting the crisp alpine air that drifted down from the mountains.',
                'She had not been to this city since childhood, yet every corner seemed to whisper memories she could not quite place. The fountain in the piazza, the wrought-iron balconies draped with wisteria, the narrow alley where cats lounged in patches of sunlight.',
                '"You must find the bookshop on Via della Scala," her grandmother had said, her voice barely above a whisper. "Ask for Signore Benedetti. He will know what to do with the manuscript."',
                'Clara unfolded the hand-drawn map tucked inside the satchel. The ink had faded in places, but the route was still legible. She set off, her footsteps echoing against the ancient stone walls.',
            ]
        },
        {
            'title': 'Rising Action',
            'paragraphs': [
                'Three weeks had passed since Clara arrived in Verona, and the mystery of the manuscript had only deepened. Signore Benedetti, the elderly bookseller, had examined the pages with trembling hands, his eyes wide behind thick spectacles.',
                '"This is written in the cipher of the Accademia degli Inchiostri," he had murmured. "A secret society of writers and scholars who operated during the Renaissance. They believed that certain stories held the power to reshape reality itself."',
                'Clara spent her days in the dim back room of the bookshop, surrounded by towers of crumbling volumes. She learned to recognize the cipher\'s patterns—the way certain letters were replaced by alchemical symbols, the hidden messages embedded in seemingly ordinary prose.',
                'Meanwhile, she noticed she was being followed. A tall figure in a grey overcoat appeared at the edge of her vision whenever she ventured out. At the cafe where she took her morning espresso. At the bridge over the Adige where she walked each evening.',
                'Benedetti warned her to be cautious. "There are others who seek the manuscript," he said. "People who understand its power far better than we do. The Accademia had enemies—and some of those enemies have very long memories."',
                'That night, Clara returned to her rented apartment to find the lock forced open and her belongings scattered across the floor. Nothing appeared to be missing—except a single page of notes she had made about the cipher.',
            ]
        },
        {
            'title': 'The Climax',
            'paragraphs': [
                'The confrontation came on a rain-soaked evening in the amphitheater. Clara had finally decoded the manuscript\'s central passage—a story within a story that described a hidden chamber beneath the Arena di Verona, where the Accademia had stored their most dangerous works.',
                'She descended through a passage concealed behind a loose stone in the amphitheater\'s lower arcade. The tunnel was narrow, the air thick with the smell of damp earth and centuries of silence. Her flashlight carved a thin beam through the darkness.',
                'The chamber was smaller than she had imagined—a circular room lined with shelves carved into the rock. Hundreds of manuscripts rested in stone alcoves, their leather bindings cracked but intact. In the center stood a stone pedestal, and on it lay a single book bound in deep crimson.',
                'As Clara reached for it, a voice echoed from the tunnel behind her. "I\'ve waited a very long time for someone to find this place." The figure in the grey overcoat stepped into the light. His face was gaunt, his eyes sharp and calculating.',
                '"My name is Marcus Adler," he said. "I am the last living descendant of the Accademia\'s founder. That manuscript your grandmother gave you—it was stolen from my family three generations ago."',
                'Clara\'s heart pounded, but she stood her ground. "My grandmother was no thief. She was a scholar who dedicated her life to preserving these works."',
                '"Preserving?" Adler laughed bitterly. "She was hiding them. Keeping them from the world. These stories were meant to be read, to be shared. That was the Accademia\'s entire purpose."',
            ]
        },
        {
            'title': 'Falling Action',
            'paragraphs': [
                'In the days that followed, Clara and Marcus reached an uneasy truce. They worked together in the underground chamber, cataloguing the manuscripts and debating their significance. Clara brought her scholarly precision; Marcus contributed his family\'s oral histories and decoded journals.',
                'They discovered that the Accademia\'s collection spanned five centuries of literature, philosophy, and natural science. Some manuscripts contained early drafts of works later attributed to famous authors. Others described scientific principles decades ahead of their time.',
                'Benedetti visited the chamber once, leaning heavily on his cane as he navigated the narrow tunnel. He wept when he saw the shelves. "I spent sixty years in that bookshop, believing these were lost forever," he said quietly.',
                'The question of what to do with the collection consumed them. Marcus wanted to donate everything to the University of Verona. Clara argued for a more careful approach—digital preservation first, then a staged release to prevent the manuscripts from being damaged by eager researchers.',
                'They compromised. The most fragile works would be digitized immediately by a conservation team Clara knew from her university days. The rest would be transferred to a climate-controlled archive at the university, with public access phased in over two years.',
                'Clara wrote to her department at Oxford, requesting a leave of absence. There was too much work to be done here, and for the first time in years, she felt the thrill of discovery that had drawn her to scholarship in the first place.',
            ]
        },
        {
            'title': 'Resolution',
            'paragraphs': [
                'One year later, Clara stood at the podium of the Verona Literary Festival, addressing an audience of scholars, journalists, and curious citizens. Behind her, a projection displayed high-resolution images of the Accademia\'s manuscripts.',
                '"What we found beneath the Arena is not merely a collection of old documents," she said. "It is a testament to the enduring belief that stories matter—that the written word has the power to illuminate, to challenge, and to transform."',
                'The audience applauded warmly. In the front row, Marcus Adler caught her eye and nodded. Beside him, Benedetti dabbed at his eyes with a handkerchief, his weathered face creased in a smile.',
                'After the lecture, Clara walked alone through the piazza. The evening air was warm, and the fountain murmured softly in the lamplight. She thought of her grandmother—of the quiet determination with which she had guarded the manuscript, and the trust she had placed in Clara to see the journey through.',
                'She reached into her satchel and touched the original manuscript, now carefully enclosed in an archival sleeve. Tomorrow it would join the rest of the collection in the university archive. But tonight, it was still hers—a bridge between the past and the future, between the stories that had been hidden and the ones yet to be told.',
                'Clara sat on the edge of the fountain and opened her notebook. She uncapped her pen and began to write.',
            ]
        },
    ]

    # Write each chapter - no separator pages in initial state
    for i, chapter in enumerate(chapters):
        # Page break before each chapter (except after title page, use section break)
        if i == 0:
            # New section for first chapter (after title page)
            new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
            new_section.page_width = Inches(8.5)
            new_section.page_height = Inches(11)
            new_section.left_margin = Inches(1)
            new_section.right_margin = Inches(1)
            new_section.top_margin = Inches(1)
            new_section.bottom_margin = Inches(1)
        else:
            # Page break before subsequent chapters
            doc.add_page_break()

        # Chapter heading
        heading = doc.add_heading(f'Chapter {i + 1}: {chapter["title"]}', level=1)

        # Add a decorative line under chapter heading
        line_para = doc.add_paragraph()
        line_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = line_para.add_run('— — —')
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.size = Pt(12)

        # Chapter body paragraphs
        for para_text in chapter['paragraphs']:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.5
            for run in p.runs:
                run.font.name = 'Georgia'
                run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
