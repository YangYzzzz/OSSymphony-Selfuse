"""
Initial Setup: 15-page Writer document with Arabic page numbering throughout.
Task ID: writer_fs_045
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_045'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_footer(section):
    """Add a footer with page number field to a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # PAGE field code: begin, instrText, end
    r1 = fp.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)

    r3 = fp.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def create_initial():
    doc = Document()

    # Set default page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Set page number type to Arabic (decimal), starting at 1
    sect_pr = section._sectPr
    pg_num_type = sect_pr.makeelement(qn('w:pgNumType'), {qn('w:fmt'): 'decimal', qn('w:start'): '1'})
    sect_pr.append(pg_num_type)

    # Add page number footer
    add_page_number_footer(section)

    # ===== PAGE 1: Title Page =====
    # Add some spacing before title
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(120)
    spacer.paragraph_format.space_after = Pt(0)

    title = doc.add_heading('Sustainable Urban Development in Metropolitan Areas', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_before = Pt(24)
    run = subtitle.add_run('A Comprehensive Analysis of Green Infrastructure\nand Smart City Initiatives')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    author = doc.add_paragraph()
    author.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_before = Pt(48)
    run = author.add_run('Dr. Elena Vasquez\nInstitute for Urban Planning and Sustainability\nMarch 2025')
    run.font.size = Pt(12)

    # Fill rest of page 1
    for _ in range(8):
        doc.add_paragraph()

    # ===== PAGE 2: Abstract =====
    doc.add_page_break()
    abstract_heading = doc.add_heading('Abstract', level=1)
    abstract_heading.paragraph_format.space_before = Pt(24)

    abstract_text = (
        "This report examines the current state of sustainable urban development across "
        "15 major metropolitan areas worldwide. Through a mixed-methods approach combining "
        "quantitative analysis of environmental metrics with qualitative assessments of "
        "community engagement, we identify key success factors for green infrastructure "
        "implementation. Our findings indicate that cities investing more than 3.5% of "
        "their municipal budget in sustainability initiatives see measurable improvements "
        "in air quality indices (average 18% reduction in PM2.5 levels), urban heat island "
        "mitigation (1.2°C average reduction), and resident satisfaction scores "
        "(23% increase over five-year periods)."
    )
    p = doc.add_paragraph(abstract_text)
    p.paragraph_format.space_after = Pt(12)

    abstract_text2 = (
        "The study further explores the role of smart city technologies in optimizing "
        "resource allocation, examining IoT sensor networks, AI-driven traffic management "
        "systems, and blockchain-based energy trading platforms. We propose a holistic "
        "framework—the Urban Sustainability Index (USI)—that integrates environmental, "
        "social, and economic indicators to guide policy decisions. Case studies from "
        "Copenhagen, Singapore, Melbourne, and Medellín demonstrate that the USI can "
        "predict long-term sustainability outcomes with 87% accuracy when calibrated "
        "against five years of historical data."
    )
    p2 = doc.add_paragraph(abstract_text2)
    p2.paragraph_format.space_after = Pt(12)

    keywords = doc.add_paragraph()
    run_k = keywords.add_run('Keywords: ')
    run_k.bold = True
    keywords.add_run('urban sustainability, green infrastructure, smart cities, '
                     'environmental metrics, urban planning, climate resilience')

    # ===== PAGE 3: Table of Contents =====
    doc.add_page_break()
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.paragraph_format.space_before = Pt(24)

    toc_entries = [
        ('1.  Introduction', '4'),
        ('2.  Literature Review', '5'),
        ('    2.1  Green Infrastructure Frameworks', '5'),
        ('    2.2  Smart City Technologies', '6'),
        ('3.  Methodology', '7'),
        ('    3.1  Data Collection', '7'),
        ('    3.2  Analysis Framework', '8'),
        ('4.  Results', '9'),
        ('    4.1  Environmental Metrics', '9'),
        ('    4.2  Social Impact Assessment', '10'),
        ('    4.3  Economic Analysis', '10'),
        ('5.  Case Studies', '11'),
        ('    5.1  Copenhagen', '11'),
        ('    5.2  Singapore', '12'),
        ('    5.3  Melbourne', '12'),
        ('    5.4  Medellín', '13'),
        ('6.  Discussion', '13'),
        ('7.  Conclusions and Recommendations', '14'),
        ('References', '15'),
    ]
    for entry, page_num in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run_entry = p.add_run(entry)
        if not entry.startswith('    '):
            run_entry.bold = True
        # Add tab and page number
        p.add_run('\t' + page_num)

    # ===== PAGES 4-15: Main Content =====

    # Page 4: Introduction
    doc.add_page_break()
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "The rapid pace of urbanization presents both unprecedented challenges and "
        "opportunities for sustainable development. By 2050, an estimated 68% of the "
        "world's population will reside in urban areas, placing enormous pressure on "
        "existing infrastructure, natural resources, and social systems. Cities already "
        "account for approximately 75% of global energy consumption and 70% of "
        "greenhouse gas emissions, making urban sustainability a critical priority "
        "for addressing climate change."
    )
    doc.add_paragraph(
        "Green infrastructure—defined as a strategically planned network of natural "
        "and semi-natural areas designed to deliver ecosystem services—has emerged as "
        "a cornerstone of sustainable urban planning. From bioswales and green roofs "
        "to urban forests and permeable pavements, these interventions offer "
        "multifunctional benefits including stormwater management, biodiversity "
        "conservation, air quality improvement, and enhanced community well-being."
    )
    doc.add_paragraph(
        "Simultaneously, smart city technologies leverage data analytics, Internet of "
        "Things (IoT) sensors, and artificial intelligence to optimize urban systems. "
        "When integrated with green infrastructure, these technologies create synergies "
        "that amplify environmental benefits while improving operational efficiency. "
        "This report examines how leading cities worldwide are combining these approaches "
        "to create more resilient, livable, and sustainable urban environments."
    )

    # Page 5: Literature Review
    doc.add_page_break()
    doc.add_heading('2. Literature Review', level=1)
    doc.add_heading('2.1 Green Infrastructure Frameworks', level=2)
    doc.add_paragraph(
        "Benedict and McMahon (2006) established the foundational framework for green "
        "infrastructure planning, emphasizing the importance of connectivity between "
        "natural areas within urban landscapes. Their work highlighted that fragmented "
        "green spaces provide fewer ecosystem services than interconnected networks. "
        "Subsequent research by Tzoulas et al. (2007) expanded this framework to "
        "include public health benefits, demonstrating correlations between urban "
        "green space access and reduced rates of cardiovascular disease, depression, "
        "and respiratory illness."
    )
    doc.add_paragraph(
        "More recently, the concept of nature-based solutions (NBS) has gained "
        "prominence in European policy contexts (European Commission, 2020). NBS "
        "extends traditional green infrastructure by emphasizing solutions that are "
        "inspired and supported by nature, addressing societal challenges while "
        "providing environmental, social, and economic co-benefits. Raymond et al. "
        "(2017) developed evaluation criteria for NBS that inform the assessment "
        "framework used in this study."
    )

    # Page 6: Smart City Technologies
    doc.add_page_break()
    doc.add_heading('2.2 Smart City Technologies', level=2)
    doc.add_paragraph(
        "The smart city paradigm, as defined by Giffinger et al. (2007), encompasses "
        "six key dimensions: smart economy, smart mobility, smart environment, smart "
        "people, smart living, and smart governance. Within the environmental dimension, "
        "sensor networks and data analytics platforms enable real-time monitoring of "
        "air quality, water resources, energy consumption, and waste generation."
    )
    doc.add_paragraph(
        "Bibri and Krogstie (2017) provided a comprehensive review of big data "
        "analytics in smart sustainable cities, identifying critical gaps in data "
        "integration across urban systems. Their findings underscore the need for "
        "interoperable platforms that can synthesize environmental sensor data with "
        "socioeconomic indicators and citizen feedback mechanisms."
    )
    doc.add_paragraph(
        "The integration of artificial intelligence in urban management has accelerated "
        "rapidly since 2020. Machine learning algorithms now optimize traffic signal "
        "timing (reducing commute times by up to 25%), predict infrastructure "
        "maintenance needs (decreasing emergency repairs by 40%), and manage distributed "
        "energy resources across smart grids (improving efficiency by 15-20%)."
    )

    # Page 7: Methodology
    doc.add_page_break()
    doc.add_heading('3. Methodology', level=1)
    doc.add_heading('3.1 Data Collection', level=2)
    doc.add_paragraph(
        "We employed a mixed-methods research design combining quantitative "
        "environmental monitoring data with qualitative stakeholder interviews. "
        "Data was collected from 15 metropolitan areas across six continents between "
        "January 2022 and December 2024. Environmental metrics included PM2.5 "
        "concentrations, urban heat island intensity, stormwater runoff volumes, "
        "biodiversity indices, and carbon sequestration rates."
    )
    doc.add_paragraph(
        "Semi-structured interviews were conducted with 127 urban planning "
        "professionals, 83 environmental scientists, 45 technology officers, and "
        "312 community members across all study sites. Interview protocols focused "
        "on implementation challenges, perceived benefits, and recommendations "
        "for improving green infrastructure and smart city integration."
    )

    # Page 8: Analysis Framework
    doc.add_page_break()
    doc.add_heading('3.2 Analysis Framework', level=2)
    doc.add_paragraph(
        "The Urban Sustainability Index (USI) was developed as a composite metric "
        "integrating 28 indicators across three dimensions: environmental quality "
        "(12 indicators), social well-being (9 indicators), and economic viability "
        "(7 indicators). Each indicator was normalized on a 0-100 scale using "
        "min-max normalization against baseline values from 2015."
    )
    doc.add_paragraph(
        "Statistical analysis employed multivariate regression models to identify "
        "significant predictors of sustainability outcomes. Principal component "
        "analysis reduced the 28 indicators to 8 orthogonal factors explaining "
        "91.3% of total variance. Hierarchical clustering grouped cities into "
        "four sustainability profiles: Leaders, Transitioning, Developing, and "
        "Emerging."
    )

    # Page 9: Results - Environmental Metrics
    doc.add_page_break()
    doc.add_heading('4. Results', level=1)
    doc.add_heading('4.1 Environmental Metrics', level=2)
    doc.add_paragraph(
        "Cities classified as 'Leaders' (Copenhagen, Singapore, Vancouver, Stockholm) "
        "demonstrated significantly lower PM2.5 concentrations (mean 8.3 µg/m³) compared "
        "to 'Emerging' cities (mean 42.7 µg/m³, p < 0.001). Green infrastructure "
        "coverage in Leader cities averaged 47% of total urban area, compared to 12% "
        "in Emerging cities. The correlation between green infrastructure coverage and "
        "air quality improvement was strong (r = 0.83, p < 0.001)."
    )
    doc.add_paragraph(
        "Urban heat island intensity showed marked variation across city profiles. "
        "Leader cities maintained an average temperature differential of 1.8°C between "
        "urban cores and surrounding rural areas, while Emerging cities exhibited "
        "differentials of up to 7.2°C. Smart irrigation systems and green roof "
        "mandates were identified as the most effective interventions, reducing "
        "surface temperatures by an average of 3.4°C in targeted areas."
    )

    # Page 10: Social Impact and Economic Analysis
    doc.add_page_break()
    doc.add_heading('4.2 Social Impact Assessment', level=2)
    doc.add_paragraph(
        "Community surveys revealed that access to green infrastructure within a "
        "10-minute walk was the strongest predictor of resident satisfaction (β = 0.67, "
        "p < 0.001). Cities with comprehensive green networks reported 23% higher "
        "satisfaction scores and 31% lower rates of reported stress-related health "
        "conditions. Equity analysis showed that Leader cities had more uniform "
        "distribution of green spaces across income quartiles (Gini coefficient 0.18 "
        "vs. 0.52 for Emerging cities)."
    )
    doc.add_heading('4.3 Economic Analysis', level=2)
    doc.add_paragraph(
        "Return on investment for green infrastructure projects averaged 4.2:1 over "
        "a 20-year period when accounting for avoided costs (flood damage, healthcare "
        "expenditures, energy savings) and increased property values. Smart city "
        "technology investments showed faster payback periods (average 3.7 years) "
        "primarily through operational efficiencies in water management, energy "
        "distribution, and waste collection."
    )

    # Page 11: Case Studies
    doc.add_page_break()
    doc.add_heading('5. Case Studies', level=1)
    doc.add_heading('5.1 Copenhagen', level=2)
    doc.add_paragraph(
        "Copenhagen's Finger Plan, updated in 2019, integrates green corridors "
        "radiating from the city center with a comprehensive cycling infrastructure "
        "network. The city's Climate Adaptation Plan allocates €1.3 billion for "
        "cloudburst management, including 300 green roofs, 60 rain gardens, and "
        "12 climate-adapted parks. IoT sensors monitor rainfall intensity in real "
        "time, automatically activating flood barriers and redirecting stormwater "
        "to retention basins. Result: 35% reduction in flood damage since 2017."
    )

    # Page 12: Singapore and Melbourne
    doc.add_page_break()
    doc.add_heading('5.2 Singapore', level=2)
    doc.add_paragraph(
        "Singapore's City in a Garden vision has transformed the city-state into "
        "one of the greenest urban environments globally. The integration of "
        "vertical gardens on 80% of new commercial buildings, combined with the "
        "ABC Waters Program (Active, Beautiful, Clean), has increased urban green "
        "coverage to 47% despite severe land constraints. Smart sensor networks "
        "monitor 850+ parks in real time."
    )
    doc.add_heading('5.3 Melbourne', level=2)
    doc.add_paragraph(
        "Melbourne's Urban Forest Strategy aims to increase canopy cover from "
        "22% to 40% by 2040. The city's innovative use of citizen science "
        "platforms—including tree email addresses that generated over 60,000 "
        "citizen messages—has created unprecedented community engagement. "
        "Machine learning models predict tree health using satellite imagery "
        "and ground-based sensors, enabling proactive maintenance."
    )

    # Page 13: Medellín and Discussion
    doc.add_page_break()
    doc.add_heading('5.4 Medellín', level=2)
    doc.add_paragraph(
        "Medellín's Green Corridors project converted 18 roads and 12 waterways "
        "into interconnected green corridors, reducing urban temperatures by 2°C "
        "and increasing biodiversity by 30%. The project combined traditional "
        "landscaping with smart monitoring systems, using AI to optimize watering "
        "schedules and track vegetation health via drone surveys."
    )
    doc.add_heading('6. Discussion', level=1)
    doc.add_paragraph(
        "Our findings confirm that integrated approaches combining green "
        "infrastructure with smart city technologies produce superior sustainability "
        "outcomes compared to isolated interventions. The USI framework reveals "
        "that the most successful cities share several characteristics: strong "
        "political commitment, dedicated funding mechanisms, cross-sector "
        "collaboration, and robust community engagement processes."
    )

    # Page 14: Conclusions
    doc.add_page_break()
    doc.add_heading('7. Conclusions and Recommendations', level=1)
    doc.add_paragraph(
        "Based on our comprehensive analysis, we offer the following recommendations "
        "for cities seeking to enhance urban sustainability:"
    )
    recommendations = [
        "Allocate a minimum of 3.5% of municipal budget to integrated sustainability initiatives",
        "Establish cross-departmental coordination bodies for green infrastructure planning",
        "Deploy IoT sensor networks for real-time environmental monitoring and adaptive management",
        "Implement equity-focused green space distribution policies targeting underserved communities",
        "Develop open data platforms that enable citizen engagement and transparency",
        "Create public-private partnership frameworks for smart city technology deployment",
        "Adopt the Urban Sustainability Index for standardized progress tracking and benchmarking",
    ]
    for rec in recommendations:
        doc.add_paragraph(rec, style='List Bullet')

    doc.add_paragraph(
        "The transition to sustainable urban environments requires sustained "
        "investment, innovative governance models, and genuine community partnerships. "
        "As urbanization accelerates, the frameworks and technologies examined in "
        "this report offer proven pathways for creating cities that are resilient, "
        "equitable, and environmentally responsible."
    )

    # Page 15: References
    doc.add_page_break()
    doc.add_heading('References', level=1)
    refs = [
        "Benedict, M. A., & McMahon, E. T. (2006). Green Infrastructure: Linking Landscapes and Communities. Island Press.",
        "Bibri, S. E., & Krogstie, J. (2017). Smart sustainable cities of the future. Sustainable Cities and Society, 31, 183-212.",
        "European Commission. (2020). Nature-Based Solutions: State of the Art in EU-Funded Projects. Publications Office.",
        "Giffinger, R., et al. (2007). Smart Cities: Ranking of European Medium-Sized Cities. Vienna UT.",
        "IPCC. (2022). Climate Change 2022: Mitigation of Climate Change. Cambridge University Press.",
        "Raymond, C. M., et al. (2017). A framework for assessing and implementing the co-benefits of NBS. Environmental Science & Policy, 77, 15-24.",
        "Tzoulas, K., et al. (2007). Promoting ecosystem and human health using green infrastructure. Landscape and Urban Planning, 81(3), 167-178.",
        "UN-Habitat. (2022). World Cities Report 2022: Envisaging the Future of Cities. United Nations.",
        "WHO. (2021). Urban Green Space Interventions and Health. World Health Organization.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
