"""
Initial Setup: Outline-numbered list with items at level 1 and level 2
Task ID: writer_lec_013
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_013'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_multilevel_numbering(doc):
    """Add a multi-level numbering definition to the document."""
    # Access or create the numbering part
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part._element

    # Create abstract numbering definition with outline levels
    abstract_num_xml = f'''
    <w:abstractNum w:abstractNumId="10" {nsdecls('w')}>
        <w:multiLevelType w:val="multilevel"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="360" w:hanging="360"/>
            </w:pPr>
            <w:rPr>
                <w:sz w:val="24"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="792" w:hanging="432"/>
            </w:pPr>
            <w:rPr>
                <w:sz w:val="24"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="2">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2.%3"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="1224" w:hanging="504"/>
            </w:pPr>
        </w:lvl>
    </w:abstractNum>
    '''
    abstract_num = parse_xml(abstract_num_xml)
    numbering_elm.insert(0, abstract_num)

    # Create concrete numbering reference
    num_xml = f'''
    <w:num w:numId="10" {nsdecls('w')}>
        <w:abstractNumId w:val="10"/>
    </w:num>
    '''
    num_elm = parse_xml(num_xml)
    numbering_elm.append(num_elm)

    return 10  # numId


def add_outline_para(doc, text, level, num_id, font_size=12, bold=False):
    """Add a paragraph with outline numbering at the given level (0-based)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold

    # Set numbering properties
    pPr = para._element.get_or_add_pPr()
    numPr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'  <w:ilvl w:val="{level}"/>'
        f'  <w:numId w:val="{num_id}"/>'
        f'</w:numPr>'
    )
    pPr.insert(0, numPr)

    return para


def create_initial():
    doc = Document()

    # Add a title
    title = doc.add_heading('Project Planning Document', level=0)

    # Add some introductory text
    intro = doc.add_paragraph(
        'This document outlines the key components of the upcoming '
        'infrastructure modernization project. Each section covers a critical '
        'aspect of the planning and execution phases.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # Force numbering part to exist by adding a dummy list paragraph and removing it
    dummy = doc.add_paragraph('', style='List Number')
    dummy._element.getparent().remove(dummy._element)

    # Add multi-level numbering
    num_id = add_multilevel_numbering(doc)

    # Level 1, Item 1: Project Overview
    add_outline_para(doc, 'Project Overview', level=0, num_id=num_id, font_size=13, bold=True)

    # Level 2 items under Project Overview
    add_outline_para(doc, 'Budget Analysis', level=1, num_id=num_id, font_size=12)

    p = doc.add_paragraph(
        'The budget analysis covers all anticipated costs including '
        'hardware procurement ($245,000), software licensing ($89,500), '
        'and consulting fees ($120,000). Total estimated budget: $454,500.'
    )
    p.paragraph_format.left_indent = Inches(0.75)
    p.paragraph_format.space_after = Pt(6)

    add_outline_para(doc, 'Risk Assessment', level=1, num_id=num_id, font_size=12)

    p = doc.add_paragraph(
        'Key risks identified include vendor dependency (high), '
        'timeline slippage due to resource constraints (medium), '
        'and integration challenges with legacy systems (high).'
    )
    p.paragraph_format.left_indent = Inches(0.75)
    p.paragraph_format.space_after = Pt(6)

    add_outline_para(doc, 'Team Roles', level=1, num_id=num_id, font_size=12)

    p = doc.add_paragraph(
        'Project Manager: Elena Rodriguez. Technical Lead: James Park. '
        'QA Lead: Priya Sharma. DevOps Engineer: Carlos Mendez. '
        'Business Analyst: Sophie Chen.'
    )
    p.paragraph_format.left_indent = Inches(0.75)
    p.paragraph_format.space_after = Pt(6)

    # Level 1, Item 2: Timeline
    add_outline_para(doc, 'Timeline', level=0, num_id=num_id, font_size=13, bold=True)

    p = doc.add_paragraph(
        'Phase 1 (Q1 2026): Requirements gathering and vendor selection. '
        'Phase 2 (Q2 2026): Infrastructure setup and initial deployment. '
        'Phase 3 (Q3 2026): Migration and testing. '
        'Phase 4 (Q4 2026): Go-live and post-launch support.'
    )
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
