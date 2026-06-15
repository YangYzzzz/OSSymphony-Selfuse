"""
Initial Setup: Writer document with default equal margins
Task ID: writer_bs_093
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_093'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default page style: equal 2.54cm margins on all sides (standard default)
    section = doc.sections[0]
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Title
    title = doc.add_heading('Quarterly Sales Performance Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Greenfield Technologies Inc.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
    run = subtitle.add_run('\nQ1 2025 - Confidential')
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph()  # spacer

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report presents the quarterly sales performance data for Greenfield Technologies Inc. '
        'covering the period from January 1 through March 31, 2025. Overall revenue increased by 12.3% '
        'compared to Q4 2024, driven primarily by strong performance in the Enterprise Solutions division '
        'and the successful launch of the CloudSync Pro platform in February.'
    )
    doc.add_paragraph(
        'Key highlights include a 28% increase in new customer acquisitions, expansion into three new '
        'regional markets (Southeast Asia, Eastern Europe, and South America), and a reduction in '
        'customer churn rate from 4.2% to 2.8%. The sales team exceeded their combined target by $1.4M, '
        'representing a 108% achievement rate against plan.'
    )

    # Regional Performance
    doc.add_heading('Regional Performance Breakdown', level=1)
    doc.add_paragraph(
        'The North American division reported $12.8M in revenue, representing 45% of total global sales. '
        'Sarah Chen, Regional VP for North America, attributes this growth to the expanded partnership '
        'program with Microsoft Azure and the onboarding of 14 new enterprise clients in the financial '
        'services sector.'
    )
    doc.add_paragraph(
        'European operations contributed $8.4M, with Germany ($2.9M), United Kingdom ($2.1M), and '
        'France ($1.6M) as the top three markets. The EMEA team, led by Marcus Johnson, successfully '
        'navigated regulatory changes related to the EU Digital Services Act while maintaining a 96% '
        'contract renewal rate.'
    )
    doc.add_paragraph(
        'Asia-Pacific revenue reached $5.7M, marking a 34% year-over-year increase. This surge was '
        'largely driven by the new Tokyo office opened in January and strategic partnerships with '
        'regional distributors in Singapore and Melbourne.'
    )

    # Product Line Analysis
    doc.add_heading('Product Line Analysis', level=1)

    # Table of product performance
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ['Product Line', 'Q1 Revenue', 'Growth %', 'Target Achievement']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['CloudSync Pro', '$4,230,000', '+42.1%', '118%'],
        ['DataVault Enterprise', '$8,150,000', '+8.7%', '103%'],
        ['SecureNet Gateway', '$6,890,000', '+15.3%', '112%'],
        ['Analytics Dashboard', '$3,420,000', '+5.2%', '97%'],
        ['API Integration Suite', '$4,210,000', '+22.8%', '109%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()  # spacer

    doc.add_paragraph(
        'The CloudSync Pro platform, launched on February 3rd, has exceeded all initial projections. '
        'With 847 new subscriptions in its first two months, it is on track to become our fastest-growing '
        'product line. The engineering team, under the direction of Dr. Priya Sharma, delivered the '
        'product two weeks ahead of schedule with zero critical defects at launch.'
    )

    # Financial Summary
    doc.add_heading('Financial Summary', level=1)
    doc.add_paragraph(
        'Total Q1 2025 revenue: $26,900,000. Operating expenses remained within budget at $18,430,000, '
        'yielding an operating margin of 31.5%. Net profit after tax stood at $5,723,000, representing '
        'a 21.3% net margin - a significant improvement over the 18.1% achieved in Q1 2024.'
    )
    doc.add_paragraph(
        'The finance team has revised the full-year forecast upward by 6%, projecting annual revenue of '
        '$112M compared to the original target of $105.7M. This revision accounts for the stronger-than-'
        'expected CloudSync Pro adoption and confirmed enterprise deals in the pipeline for Q2.'
    )

    # Recommendations
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph(
        'Based on Q1 performance data, the leadership team recommends the following strategic actions '
        'for Q2 2025:', style='List Bullet'
    )
    doc.add_paragraph(
        'Increase investment in CloudSync Pro marketing by $500K to capitalize on current momentum',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Accelerate hiring of 12 additional sales representatives for the Asia-Pacific region',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Initiate preliminary discussions for potential acquisition of CompetitorX analytics platform',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Launch customer success program targeting the 23 accounts identified as at-risk for churn',
        style='List Bullet'
    )

    # Footer note
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_para.add_run('Prepared by the Office of the CFO - Distribution: Board Members Only')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
