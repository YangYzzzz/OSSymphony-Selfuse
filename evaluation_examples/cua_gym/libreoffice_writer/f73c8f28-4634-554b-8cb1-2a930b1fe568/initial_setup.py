"""
Initial Setup: Create Quarterly_Financials.docx with title page and 2 blank pages
Task ID: writer_pd_028
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_028'
OUTPUT = f'{WORKDIR}/Quarterly_Financials.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ---- Page 1: Title Page ----
    # Add some vertical spacing
    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.space_before = Pt(0)

    title = doc.add_heading('Q1 2026 Financial Summary', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by Finance Department')
    run.font.size = Pt(14)
    run.font.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = date_para.add_run('Report Date: March 31, 2026')
    run2.font.size = Pt(12)

    # Page break to page 2
    doc.add_page_break()

    # ---- Page 2: Empty (placeholder for Income Statement & Balance Sheet) ----
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)

    # Page break to page 3
    doc.add_page_break()

    # ---- Page 3: Empty (placeholder for Cash Flow Statement) ----
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
