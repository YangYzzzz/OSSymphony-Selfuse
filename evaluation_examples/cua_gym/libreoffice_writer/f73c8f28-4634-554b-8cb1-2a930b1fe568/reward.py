"""
Reward Script: Financial summary with three formatted tables in Quarterly_Financials.docx
Task ID: writer_pd_028
Domain: libreoffice_writer
Scoring:
  Component 1: Three distinct tables exist (0.25)
  Component 2: Income Statement table has required rows (0.20)
  Component 3: Balance Sheet has Assets/Liabilities/Equity sections (0.20)
  Component 4: Cash Flow covers operating/investing/financing (0.10)
  Component 5: Dark header rows with white bold text (0.10)
  Component 6: Right-aligned currency columns with dollar signs (0.10)
  Component 7: Bold totals/subtotals rows (0.05)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_028'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits via Ctrl+S."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def get_all_cell_texts(table):
    """Get all cell texts from a table as a flat list."""
    texts = []
    for row in table.rows:
        for cell in row.cells:
            texts.append(cell.text.strip())
    return texts


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document
    from docx.shared import RGBColor

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    tables = doc.tables

    # Component 1: Three distinct tables exist (0.25 points)
    # Initial has 0 tables, golden has 3
    try:
        num_tables = len(tables)
        if num_tables >= 3:
            print(f"PASS: Component 1 — {num_tables} tables found (>= 3) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Expected >= 3 tables, found {num_tables}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    if len(tables) < 3:
        # Cannot proceed with detailed checks if tables don't exist
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Component 2: Income Statement table has required rows (0.20 points)
    # Check that the first table has rows containing: Revenue, COGS/Cost of Goods,
    # Gross Profit, Operating Expenses, Net Income
    try:
        t0_texts = [row.cells[0].text.strip().lower() for row in tables[0].rows]
        t0_full = ' '.join(t0_texts)
        required_income_items = [
            ('revenue', 'Revenue'),
            ('cost of goods', 'Cost of Goods Sold'),
            ('gross profit', 'Gross Profit'),
            ('operating expense', 'Operating Expenses'),
            ('net income', 'Net Income'),
        ]
        found_count = 0
        for keyword, label in required_income_items:
            if keyword in t0_full:
                found_count += 1
            else:
                print(f"  MISS: Income Statement missing '{label}'")

        # Also check that it has at least 2 quarter/data columns (>= 3 total cols)
        has_data_cols = len(tables[0].columns) >= 3

        if found_count == len(required_income_items) and has_data_cols:
            print(f"PASS: Component 2 — Income Statement has all {found_count} required items and {len(tables[0].columns)} cols (0.20 pts)")
            total_score += 0.20
        elif found_count >= 3:
            partial = 0.10
            print(f"PARTIAL: Component 2 — Income Statement has {found_count}/{len(required_income_items)} items ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Income Statement has only {found_count}/{len(required_income_items)} items")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Balance Sheet has Assets/Liabilities/Equity sections (0.20 points)
    try:
        t1_texts = [row.cells[0].text.strip().lower() for row in tables[1].rows]
        t1_full = ' '.join(t1_texts)
        balance_sections = [
            ('assets', 'Assets section'),
            ('liabilities', 'Liabilities section'),
            ('equity', 'Equity section'),
            ('total assets', 'Total Assets subtotal'),
            ('total liabilities', 'Total Liabilities subtotal'),
            ('total equity', 'Total Equity subtotal'),
        ]
        found_sections = 0
        for keyword, label in balance_sections:
            if keyword in t1_full:
                found_sections += 1
            else:
                print(f"  MISS: Balance Sheet missing '{label}'")

        if found_sections == len(balance_sections):
            print(f"PASS: Component 3 — Balance Sheet has all {found_sections} required sections/subtotals (0.20 pts)")
            total_score += 0.20
        elif found_sections >= 4:
            partial = 0.10
            print(f"PARTIAL: Component 3 — Balance Sheet has {found_sections}/{len(balance_sections)} sections ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Balance Sheet has only {found_sections}/{len(balance_sections)} sections")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Cash Flow covers operating/investing/financing (0.10 points)
    try:
        t2_texts = [row.cells[0].text.strip().lower() for row in tables[2].rows]
        t2_full = ' '.join(t2_texts)
        cf_sections = [
            ('operating', 'Operating Activities'),
            ('investing', 'Investing Activities'),
            ('financing', 'Financing Activities'),
        ]
        found_cf = 0
        for keyword, label in cf_sections:
            if keyword in t2_full:
                found_cf += 1
            else:
                print(f"  MISS: Cash Flow missing '{label}'")

        if found_cf == len(cf_sections):
            print(f"PASS: Component 4 — Cash Flow has all {found_cf} activity sections (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — Cash Flow has only {found_cf}/{len(cf_sections)} sections")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Dark header rows with white bold text (0.10 points)
    # Check that header rows (row 0) of all 3 tables have dark shading and white bold text
    try:
        ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        dark_header_count = 0
        for ti in range(3):
            header_row = tables[ti].rows[0]
            # Check shading on first cell
            shd_elem = header_row.cells[0]._element.find(f'.//{{{ns_w}}}shd')
            has_dark_shading = False
            if shd_elem is not None:
                fill_val = shd_elem.get(f'{{{ns_w}}}fill', '')
                # Dark colors have low RGB values; check if fill is a dark hex color
                if fill_val and fill_val != 'auto':
                    try:
                        r_val = int(fill_val[0:2], 16)
                        g_val = int(fill_val[2:4], 16)
                        b_val = int(fill_val[4:6], 16)
                        # Dark if average < 128
                        if (r_val + g_val + b_val) / 3 < 128:
                            has_dark_shading = True
                    except (ValueError, IndexError):
                        pass

            # Check bold white text in header
            has_bold_white = False
            for cell in header_row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.bold and run.font.color and run.font.color.rgb:
                            rgb = run.font.color.rgb
                            # RGBColor is indexable: rgb[0]=R, rgb[1]=G, rgb[2]=B (ints 0-255)
                            if rgb[0] > 200 and rgb[1] > 200 and rgb[2] > 200:
                                has_bold_white = True
                                break

            if has_dark_shading and has_bold_white:
                dark_header_count += 1

        if dark_header_count == 3:
            print(f"PASS: Component 5 — All 3 tables have dark header rows with white bold text (0.10 pts)")
            total_score += 0.10
        elif dark_header_count >= 1:
            partial = round(0.10 * dark_header_count / 3, 2)
            print(f"PARTIAL: Component 5 — {dark_header_count}/3 tables have dark headers ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — No tables have dark header rows with white bold text")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Right-aligned currency columns with dollar signs (0.10 points)
    # Check that numeric/currency cells in data rows are right-aligned and contain $
    try:
        tables_with_aligned_currency = 0
        for ti in range(3):
            table = tables[ti]
            aligned_dollar_cells = 0
            total_dollar_cells = 0
            for ri in range(1, len(table.rows)):
                for ci in range(1, len(table.columns)):
                    cell_text = table.rows[ri].cells[ci].text.strip()
                    if '$' in cell_text or (cell_text.startswith('(') and '$' in cell_text):
                        total_dollar_cells += 1
                        # Check alignment
                        for para in table.rows[ri].cells[ci].paragraphs:
                            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                            if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                                aligned_dollar_cells += 1
                                break

            if total_dollar_cells > 0 and aligned_dollar_cells >= total_dollar_cells * 0.7:
                tables_with_aligned_currency += 1

        if tables_with_aligned_currency == 3:
            print(f"PASS: Component 6 — All 3 tables have right-aligned currency columns (0.10 pts)")
            total_score += 0.10
        elif tables_with_aligned_currency >= 1:
            partial = round(0.10 * tables_with_aligned_currency / 3, 2)
            print(f"PARTIAL: Component 6 — {tables_with_aligned_currency}/3 tables have aligned currency ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 6 — No tables have right-aligned currency columns")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Bold totals/subtotals rows (0.05 points)
    # Check that rows with "Total" or key summary labels have bold formatting
    try:
        bold_total_count = 0
        total_total_count = 0
        for ti in range(3):
            for row in tables[ti].rows:
                first_text = row.cells[0].text.strip()
                if 'Total' in first_text or first_text in ('Gross Profit', 'Net Income', 'Operating Income'):
                    total_total_count += 1
                    row_has_bold = False
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.font.bold:
                                    row_has_bold = True
                                    break
                            if row_has_bold:
                                break
                        if row_has_bold:
                            break
                    if row_has_bold:
                        bold_total_count += 1

        if total_total_count > 0 and bold_total_count >= total_total_count * 0.6:
            print(f"PASS: Component 7 — {bold_total_count}/{total_total_count} total/subtotal rows are bold (0.05 pts)")
            total_score += 0.05
        elif bold_total_count > 0:
            print(f"PARTIAL: Component 7 — Only {bold_total_count}/{total_total_count} total rows are bold (0.02 pts)")
            total_score += 0.02
        else:
            print(f"FAIL: Component 7 — No total/subtotal rows have bold formatting")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/Quarterly_Financials.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
