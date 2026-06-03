"""
Initial Setup: Create a Writer document for AutoText entry task
Task ID: writer_acad_089
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_089'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Title
    title = doc.add_heading('Analyzing Urban Heat Island Effects Using Machine Learning Approaches', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Authors
    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = authors.add_run('Elena Rodriguez, James Park, and Aisha Patel')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    affil = doc.add_paragraph()
    affil.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = affil.add_run('Department of Environmental Science, Pacific Northwest University')
    run.font.size = Pt(10)
    run.italic = True

    # Abstract
    doc.add_heading('Abstract', level=2)
    doc.add_paragraph(
        'Urban heat islands (UHIs) represent a significant environmental challenge in '
        'metropolitan areas worldwide. This study employs gradient-boosted decision trees '
        'and convolutional neural networks to predict surface temperature variations across '
        '47 urban centers in the United States. Our analysis of satellite thermal imagery '
        'from 2018 to 2023 reveals that impervious surface coverage and vegetation density '
        'are the strongest predictors of localized temperature anomalies, accounting for '
        '68.3% of observed variance (R² = 0.683, p < 0.001). The proposed ensemble model '
        'achieves a mean absolute error of 1.2°C on the held-out test set, outperforming '
        'traditional regression approaches by 34%.'
    )

    # Introduction
    doc.add_heading('1. Introduction', level=2)
    doc.add_paragraph(
        'The urban heat island effect—wherein metropolitan areas experience elevated '
        'temperatures relative to surrounding rural regions—has been extensively documented '
        'since Manley (1958) first quantified the phenomenon in London. Contemporary '
        'estimates suggest that large cities may be 1–3°C warmer than adjacent non-urban '
        'areas during daytime and up to 12°C warmer at night (Oke et al., 2017). These '
        'temperature differentials carry profound implications for public health, energy '
        'consumption, and ecosystem functioning.'
    )
    doc.add_paragraph(
        'Recent advances in machine learning have opened new avenues for modeling complex '
        'spatiotemporal phenomena. Random forests, support vector machines, and deep learning '
        'architectures have all demonstrated utility in environmental prediction tasks '
        '(Reichstein et al., 2019). However, comparatively few studies have applied these '
        'methods specifically to UHI prediction at the continental scale. This gap motivates '
        'the present investigation.'
    )

    # Methods
    doc.add_heading('2. Methods', level=2)
    doc.add_heading('2.1 Data Collection', level=3)
    doc.add_paragraph(
        'We obtained Landsat 8 thermal infrared (Band 10, 10.6–11.2 µm) imagery for 47 '
        'U.S. metropolitan statistical areas from the USGS Earth Explorer platform. Images '
        'were selected from June through August for years 2018–2023, yielding a total of '
        '1,128 cloud-free scenes. Land use/land cover data were sourced from the National '
        'Land Cover Database (NLCD) 2021 edition, providing 30-meter resolution classification '
        'of impervious surface, vegetation, water body, and bare soil fractions.'
    )

    doc.add_heading('2.2 Feature Engineering', level=3)
    doc.add_paragraph(
        'For each 250-meter grid cell within the study domains, we computed the following '
        'predictor variables: (a) impervious surface fraction (ISF), (b) normalized difference '
        'vegetation index (NDVI), (c) distance to nearest water body, (d) building height '
        'from LIDAR-derived digital surface models, (e) sky view factor, and (f) population '
        'density from the American Community Survey 2022 estimates. All features were '
        'standardized to zero mean and unit variance prior to model fitting.'
    )

    # Results table
    doc.add_heading('3. Preliminary Results', level=2)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Model', 'R²', 'MAE (°C)', 'RMSE (°C)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True

    data = [
        ['Linear Regression', '0.421', '2.14', '2.87'],
        ['Random Forest', '0.589', '1.73', '2.31'],
        ['Gradient Boosting (XGBoost)', '0.651', '1.38', '1.92'],
        ['CNN (ResNet-18)', '0.667', '1.29', '1.78'],
        ['Ensemble (XGBoost + CNN)', '0.683', '1.21', '1.65'],
    ]
    for r_idx, row_data in enumerate(data, 1):
        for c_idx, val in enumerate(row_data):
            table.cell(r_idx, c_idx).text = val

    doc.add_paragraph()  # spacing

    # Discussion
    doc.add_heading('4. Discussion', level=2)
    doc.add_paragraph(
        'The ensemble model combining gradient-boosted trees with convolutional features '
        'consistently outperformed individual architectures across all evaluation metrics. '
        'Shapley additive explanation (SHAP) analysis reveals that impervious surface '
        'fraction contributes the largest marginal effect on predicted temperature (mean '
        '|SHAP| = 0.42°C), followed by NDVI (mean |SHAP| = 0.31°C) and building height '
        '(mean |SHAP| = 0.18°C). These findings corroborate the physical mechanisms '
        'underlying urban thermal dynamics: dark, impervious materials absorb and re-emit '
        'solar radiation more efficiently than vegetated surfaces, while tall buildings '
        'reduce radiative cooling by limiting sky exposure.'
    )

    # Acknowledgments placeholder (without the NSF grant text)
    doc.add_heading('Acknowledgments', level=2)
    doc.add_paragraph(
        'The authors thank Dr. Wei Zhang for helpful discussions on feature selection '
        'methodology and the Pacific Northwest University High-Performance Computing Center '
        'for providing computational resources.'
    )

    # References
    doc.add_heading('References', level=2)
    refs = [
        'Manley, G. (1958). On the frequency of snowfall in metropolitan England. Quarterly Journal of the Royal Meteorological Society, 84(359), 70–72.',
        'Oke, T. R., Mills, G., Christen, A., & Voogt, J. A. (2017). Urban Climates. Cambridge University Press.',
        'Reichstein, M., Camps-Valls, G., Stevens, B., et al. (2019). Deep learning and process understanding for data-driven Earth system science. Nature, 566(7743), 195–204.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Ensure no custom AutoText entries exist (clear mytexts.bau to empty state)
    autotext_dir = '/home/user/.config/libreoffice/4/user/autotext'
    bau_path = os.path.join(autotext_dir, 'mytexts.bau')
    if os.path.exists(bau_path):
        # Recreate empty bau file
        import zipfile
        os.makedirs(autotext_dir, exist_ok=True)
        with zipfile.ZipFile(bau_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr('mimetype', '')
            zf.writestr('BlockList.xml',
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<block-list:block-list xmlns:block-list="http://openoffice.org/2001/block-list"'
                ' block-list:list-name="My AutoText"/>\n')
            zf.writestr('META-INF/manifest.xml',
                '<manifest:manifest xmlns:manifest="http://openoffice.org/2001/manifest">'
                '<manifest:file-entry manifest:media-type="" manifest:full-path="/"/>'
                '<manifest:file-entry manifest:media-type="" manifest:full-path="META-INF/"/>'
                '<manifest:file-entry manifest:media-type="text/xml" manifest:full-path="BlockList.xml"/>'
                '</manifest:manifest>')
    print('AutoText cleared: no custom entries exist')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
