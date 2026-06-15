"""
FINAL REWARD SCRIPT - SUCCESS
Task: My Writer file is covered in spots where I’ve accidentally typed two consecutive spaces instead of one. Before I hand it off, I need a quick way to scan the entire document and replace every exact instance of "  " (two spaces) with a single " " (one space). How do I pull that off in LibreOffice Writer?
Generated: 2025-09-10 13:49:59
Status: success
Model: azure-o3
Total Steps: 2
"""

import os
import re
from docx import Document


def _gather_all_text(doc: Document) -> str:
    """Extract every visible text fragment from a DOCX document, including
    main body paragraphs, tables (recursively), headers and footers."""

    texts = []

    def extract_table(table):
        for row in table.rows:
            for cell in row.cells:
                # Paragraphs in the current cell
                for p in cell.paragraphs:
                    texts.append(p.text)
                # Nested tables inside this cell (can be multi-level)
                for nested in cell.tables:
                    extract_table(nested)

    # --- Body paragraphs & tables ---
    for para in doc.paragraphs:
        texts.append(para.text)
    for tbl in doc.tables:
        extract_table(tbl)

    # --- Headers & footers ---
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            for para in part.paragraphs:
                texts.append(para.text)
            for tbl in part.tables:
                extract_table(tbl)

    return "\n".join(texts)


def verify_no_consecutive_spaces(file_path: str) -> float:
    """Verify that the document contains no occurrences of two (or more)
    consecutive space characters.  Returns a progressive score 0.0–1.0.
    One exact task requirement ⇒ score 1.0 when fully satisfied.
    A graceful degradation is applied if violations are found.
    """

    # 0) Basic pre-checks (no points awarded!)
    if not os.path.exists(file_path):
        print("✗ File not found:", file_path)
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as exc:
        print("✗ Unable to load DOCX:", exc)
        return 0.0

    # 1) Collect every visible text fragment
    full_text = _gather_all_text(doc)
    total_chars = len(full_text)
    if total_chars == 0:
        print("✗ Document appears to contain no text – cannot verify task")
        return 0.0

    # 2) Detect consecutive spaces (two or more)
    consecutive_matches = re.findall(r" {2,}", full_text)
    violations = len(consecutive_matches)
    print(f"Detected {violations} occurrence(s) of consecutive spaces in the document.")

    # 3) Progressive scoring – perfect only if zero violations.
    #    We linearly reduce the score with a cap so that huge numbers don’t
    #    drop it below 0.0 in edge cases.
    if violations == 0:
        score = 1.0
    else:
        # Treat 30+ violations as complete failure (score 0). Fewer violations
        # degrade the score proportionally.
        penalty_ratio = min(violations, 30) / 30.0  # range 0-1
        score = max(0.0, 1.0 - penalty_ratio)

    print(f"Computed reward score: {score}")
    return score


# ---------------------------------------------------------------------------
# MAIN EXECUTION (called when the script is executed by the evaluation system)
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    DOC_PATH = "/home/user/my_writer_file_is_covered_in_spots_where_ive_accidentally_typed_two_consecutive_spaces_instead_of_on.docx"

    final_reward = verify_no_consecutive_spaces(DOC_PATH)
    print(f"REWARD: {final_reward}")

