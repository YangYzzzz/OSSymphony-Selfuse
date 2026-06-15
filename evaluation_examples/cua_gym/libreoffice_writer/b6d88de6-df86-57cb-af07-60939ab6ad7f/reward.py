"""
Reward Script: API Documentation Template with Custom Styles
Task ID: writer_tech_055
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Custom paragraph styles defined (CodeBlock, NoteBox, ParameterTable)
  Component 2 (0.15): Custom character styles defined (InlineCode, APIMethod)
  Component 3 (0.25): CodeBlock style applied to code block paragraphs
  Component 4 (0.10): NoteBox style applied to note paragraphs
  Component 5 (0.15): Cover page has placeholder fields ([VERSION_NUMBER], [AUTHOR_NAME], etc.)
  Component 6 (0.10): APIMethod or InlineCode character styles applied to runs
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_055'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Custom paragraph styles (CodeBlock, NoteBox, ParameterTable) exist (0.25 pts)
    # These styles do NOT exist in the initial file; they are task-introduced.
    try:
        style_names = [s.name for s in doc.styles if s.type is not None]
        required_para_styles = ['CodeBlock', 'NoteBox', 'ParameterTable']
        found_para = [s for s in required_para_styles if s in style_names]
        if len(found_para) == 3:
            print(f"PASS: Component 1 - All 3 custom paragraph styles found: {found_para} (0.25 pts)")
            total_score += 0.25
        elif len(found_para) >= 2:
            partial = round(0.25 * len(found_para) / 3, 2)
            print(f"PARTIAL: Component 1 - {len(found_para)}/3 custom paragraph styles found: {found_para} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 - Expected CodeBlock, NoteBox, ParameterTable styles, found: {found_para}")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Custom character styles (InlineCode, APIMethod) exist (0.15 pts)
    # These styles do NOT exist in the initial file.
    try:
        required_char_styles = ['InlineCode', 'APIMethod']
        found_char = [s for s in required_char_styles if s in style_names]
        if len(found_char) == 2:
            print(f"PASS: Component 2 - Both custom character styles found: {found_char} (0.15 pts)")
            total_score += 0.15
        elif len(found_char) == 1:
            print(f"PARTIAL: Component 2 - 1/2 custom character styles found: {found_char} (0.075 pts)")
            total_score += 0.075
        else:
            print(f"FAIL: Component 2 - Expected InlineCode, APIMethod styles, found: {found_char}")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: CodeBlock style is actually applied to paragraphs (0.25 pts)
    # In initial file: 0 paragraphs use CodeBlock. In golden: 25.
    try:
        codeblock_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'CodeBlock')
        if codeblock_count >= 10:
            print(f"PASS: Component 3 - {codeblock_count} paragraphs use CodeBlock style (0.25 pts)")
            total_score += 0.25
        elif codeblock_count >= 3:
            partial = round(0.25 * min(codeblock_count, 10) / 10, 2)
            print(f"PARTIAL: Component 3 - {codeblock_count} paragraphs use CodeBlock style ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 - Expected >=10 paragraphs with CodeBlock style, found: {codeblock_count}")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: NoteBox style is applied to note paragraphs (0.10 pts)
    # In initial file: 0 paragraphs use NoteBox. In golden: 2.
    try:
        notebox_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'NoteBox')
        if notebox_count >= 1:
            print(f"PASS: Component 4 - {notebox_count} paragraphs use NoteBox style (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 - Expected >=1 paragraphs with NoteBox style, found: {notebox_count}")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Cover page has placeholder fields (0.15 pts)
    # Initial file has concrete values ("Version 3.2.1", "Last Updated: March 2026").
    # Golden file has placeholders: [VERSION_NUMBER], [LAST_UPDATED_DATE], [AUTHOR_NAME], [DEPARTMENT].
    try:
        placeholders = ['[VERSION_NUMBER]', '[LAST_UPDATED_DATE]', '[AUTHOR_NAME]', '[DEPARTMENT]']
        # Check first 10 paragraphs for placeholders
        first_paras_text = ' '.join(p.text for p in doc.paragraphs[:10])
        found_placeholders = [ph for ph in placeholders if ph in first_paras_text]
        if len(found_placeholders) >= 3:
            print(f"PASS: Component 5 - {len(found_placeholders)}/4 placeholder fields found in cover page (0.15 pts)")
            total_score += 0.15
        elif len(found_placeholders) >= 1:
            partial = round(0.15 * len(found_placeholders) / 3, 2)
            print(f"PARTIAL: Component 5 - {len(found_placeholders)}/4 placeholder fields found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 - No placeholder fields found in cover page. Text: {first_paras_text[:200]}")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: Character styles (InlineCode/APIMethod) applied to runs (0.10 pts)
    # In initial file: 0 runs use these styles. In golden: InlineCode=1, APIMethod=3.
    try:
        inline_code_runs = 0
        api_method_runs = 0
        for p in doc.paragraphs:
            for run in p.runs:
                if run.style and run.style.name == 'InlineCode':
                    inline_code_runs += 1
                if run.style and run.style.name == 'APIMethod':
                    api_method_runs += 1
        char_style_runs = inline_code_runs + api_method_runs
        if char_style_runs >= 2:
            print(f"PASS: Component 6 - {char_style_runs} runs use custom character styles (InlineCode={inline_code_runs}, APIMethod={api_method_runs}) (0.10 pts)")
            total_score += 0.10
        elif char_style_runs >= 1:
            print(f"PARTIAL: Component 6 - {char_style_runs} runs use custom character styles (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 - No runs use InlineCode or APIMethod character styles")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
