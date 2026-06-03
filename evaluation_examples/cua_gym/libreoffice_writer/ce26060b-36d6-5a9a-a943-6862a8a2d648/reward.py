"""
Reward Script: Running header with chapter title field
Task ID: writer_rd_059
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.40): STYLEREF "Heading 1" field code present in header
  - Component 2 (0.25): Right-aligned tab stop in header paragraph
  - Component 3 (0.20): "Research Paper" text retained on left side of header
  - Component 4 (0.15): Tab character present separating left text from chapter field
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_059'


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Get the first section header (the document has one section)
    try:
        section = doc.sections[0]
        header = section.header
        if not header.paragraphs:
            print("FAIL: No paragraphs in header")
            print(f"\nScore: {total_score}/1.0")
            print(f"REWARD: {total_score}")
            return total_score
        header_para = header.paragraphs[0]
    except Exception as e:
        print(f"CRITICAL: Cannot access header: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: STYLEREF "Heading 1" field code present in header (0.40 points)
    # This is the core task requirement — inserting a chapter reference field.
    # Initial state has NO field codes in header. Golden has STYLEREF "Heading 1".
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        has_styleref = False
        references_heading1 = False

        # Search across all header paragraphs for field codes
        for hp in header.paragraphs:
            fld_chars = hp._element.findall('.//w:fldChar', ns)
            instrs = hp._element.findall('.//w:instrText', ns)
            if fld_chars and instrs:
                for instr in instrs:
                    instr_text = (instr.text or '').strip().upper()
                    if 'STYLEREF' in instr_text:
                        has_styleref = True
                        # Check if it references Heading 1 (the chapter style)
                        if 'HEADING 1' in instr_text or 'HEADING1' in instr_text:
                            references_heading1 = True

        if has_styleref and references_heading1:
            print(f"PASS: Component 1 — STYLEREF 'Heading 1' field found in header (0.40 pts)")
            total_score += 0.40
        elif has_styleref:
            # Has STYLEREF but not referencing Heading 1 — partial credit
            print(f"PARTIAL: Component 1 — STYLEREF field found but not referencing 'Heading 1' (0.20 pts)")
            total_score += 0.20
        else:
            # Check for any chapter-related field (CHAPTER, CHAPTERREF, etc.)
            any_chapter_field = False
            for hp in header.paragraphs:
                instrs = hp._element.findall('.//w:instrText', ns)
                for instr in instrs:
                    instr_text = (instr.text or '').strip().upper()
                    if 'CHAPTER' in instr_text:
                        any_chapter_field = True
            if any_chapter_field:
                print(f"PARTIAL: Component 1 — Chapter field found but not STYLEREF (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 1 — No STYLEREF or chapter field found in header")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Right-aligned tab stop in header paragraph (0.25 points)
    # Initial state has NO tab stops. Golden has a RIGHT tab stop for positioning the chapter on the right.
    try:
        from docx.enum.text import WD_TAB_ALIGNMENT
        has_right_tab = False
        for hp in header.paragraphs:
            for ts in hp.paragraph_format.tab_stops:
                # Filter out default/CLEAR stops
                if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
                    continue
                if ts.alignment == WD_TAB_ALIGNMENT.RIGHT:
                    has_right_tab = True
                    break
            # Also accept right-aligned paragraph as alternative
            if hp.paragraph_format.alignment is not None:
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                if hp.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                    has_right_tab = True

        if has_right_tab:
            print(f"PASS: Component 2 — Right-aligned tab stop found in header (0.25 pts)")
            total_score += 0.25
        else:
            # Check if there's a CENTER tab stop (partial — not right but at least positioned)
            has_center_tab = False
            for hp in header.paragraphs:
                for ts in hp.paragraph_format.tab_stops:
                    if ts.alignment == WD_TAB_ALIGNMENT.CENTER:
                        has_center_tab = True
            if has_center_tab:
                print(f"PARTIAL: Component 2 — Center tab stop found but not RIGHT (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 2 — No right-aligned tab stop in header")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: "Research Paper" text retained on left side of header (0.20 points)
    # Both initial and golden have "Research Paper" — but this is a COMPOUND check:
    # "Research Paper" must coexist with the chapter field. If there's no field (initial),
    # having just "Research Paper" alone should NOT earn points.
    try:
        full_header_text = ''.join(hp.text for hp in header.paragraphs)
        has_research_paper = 'Research Paper' in full_header_text

        # Compound condition: "Research Paper" is present AND a field code exists
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        has_any_field = False
        for hp in header.paragraphs:
            fld_chars = hp._element.findall('.//w:fldChar', ns)
            if fld_chars:
                has_any_field = True
                break

        if has_research_paper and has_any_field:
            print(f"PASS: Component 3 — 'Research Paper' retained alongside chapter field (0.20 pts)")
            total_score += 0.20
        elif has_research_paper and not has_any_field:
            print(f"FAIL: Component 3 — 'Research Paper' exists but no chapter field (precondition only)")
        elif not has_research_paper and has_any_field:
            print(f"FAIL: Component 3 — Chapter field exists but 'Research Paper' was removed")
        else:
            print(f"FAIL: Component 3 — Neither 'Research Paper' nor chapter field found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Tab character separating left text from chapter field (0.15 points)
    # Initial has no tab in header. Golden uses tab to position "Research Paper" left and chapter right.
    # Compound check: tab must exist AND field must exist (otherwise just a tab in static text).
    try:
        has_tab_in_header = False
        for hp in header.paragraphs:
            for run in hp.runs:
                if run.text and '\t' in run.text:
                    has_tab_in_header = True
                    break
            if has_tab_in_header:
                break

        if has_tab_in_header and has_any_field:
            print(f"PASS: Component 4 — Tab character found separating content in header (0.15 pts)")
            total_score += 0.15
        elif has_tab_in_header:
            print(f"FAIL: Component 4 — Tab present but no field code (not meaningful separation)")
        else:
            print(f"FAIL: Component 4 — No tab character in header")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
