"""
Initial Setup: Set up a running header with chapter title field reference
Task ID: writer_rd_059
Domain: libreoffice_writer

Creates a 6-chapter research paper document with a static header "Research Paper".
The header does NOT contain any chapter field — that is the task for the agent.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_059'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Header: static "Research Paper" text only ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    run = hp.add_run("Research Paper")
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    # --- Title page content ---
    title = doc.add_heading("A Comprehensive Study on Machine Learning Applications in Healthcare", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = subtitle.add_run("Dr. Elena Vasquez, Dr. James Thornton, Dr. Mei-Ling Wu")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    affil = doc.add_paragraph()
    affil.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r2 = affil.add_run("Stanford University School of Medicine\nMarch 2025")
    r2.font.size = Pt(11)
    r2.font.name = "Times New Roman"

    doc.add_page_break()

    # --- Chapter 1: Introduction ---
    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_paragraph(
        "The intersection of machine learning and healthcare has emerged as one of the most "
        "promising frontiers in modern medicine. Over the past decade, advances in computational "
        "methods have enabled researchers to analyze vast datasets of patient records, genomic "
        "sequences, and medical imaging with unprecedented accuracy and speed."
    )
    doc.add_paragraph(
        "This paper presents a comprehensive survey of machine learning applications across "
        "multiple healthcare domains, including diagnostic imaging, drug discovery, personalized "
        "treatment planning, and epidemiological modeling. We examine both the theoretical "
        "foundations and practical implementations that have shaped the current landscape."
    )
    doc.add_paragraph(
        "Our analysis draws on data from 147 peer-reviewed studies published between 2018 and "
        "2024, encompassing clinical trials at 23 major medical institutions across North America, "
        "Europe, and East Asia. The findings suggest that ML-driven diagnostics achieve an average "
        "sensitivity of 94.3% compared to 87.1% for traditional methods."
    )

    doc.add_page_break()

    # --- Chapter 2: Literature Review ---
    doc.add_heading("Chapter 2: Literature Review", level=1)
    doc.add_paragraph(
        "Early applications of artificial intelligence in medicine date back to the 1970s with "
        "expert systems such as MYCIN, which could diagnose bacterial infections and recommend "
        "antibiotics. However, these rule-based systems were limited by the need for manually "
        "encoded domain knowledge and their inability to generalize across conditions."
    )
    doc.add_paragraph(
        "The deep learning revolution, catalyzed by Krizhevsky et al. (2012), fundamentally "
        "transformed the field. Convolutional neural networks demonstrated superhuman performance "
        "in image classification, which quickly found applications in radiology and pathology. "
        "Esteva et al. (2017) showed that a CNN trained on 129,450 clinical images could classify "
        "skin cancer with accuracy comparable to board-certified dermatologists."
    )
    doc.add_paragraph(
        "More recently, transformer-based architectures have enabled breakthroughs in natural "
        "language processing for clinical text. Models like Med-PaLM (Singhal et al., 2023) "
        "have achieved expert-level performance on medical question answering benchmarks, raising "
        "both excitement and concerns about the role of AI in clinical decision-making."
    )

    doc.add_page_break()

    # --- Chapter 3: Methodology ---
    doc.add_heading("Chapter 3: Methodology", level=1)
    doc.add_paragraph(
        "We employed a mixed-methods approach combining quantitative meta-analysis with qualitative "
        "case studies. Our systematic review followed the PRISMA guidelines, with initial screening "
        "of 2,340 articles from PubMed, IEEE Xplore, and the ACM Digital Library. After applying "
        "inclusion and exclusion criteria, 147 studies were selected for detailed analysis."
    )
    doc.add_paragraph(
        "For the quantitative component, we extracted performance metrics including accuracy, "
        "sensitivity, specificity, and area under the ROC curve (AUC-ROC) from each study. "
        "Statistical heterogeneity was assessed using the I-squared statistic, and random-effects "
        "models were applied where significant heterogeneity was detected (I-squared > 50%)."
    )
    doc.add_paragraph(
        "The qualitative component involved semi-structured interviews with 34 clinicians and "
        "12 ML researchers at five academic medical centers: Stanford Medical Center, Johns Hopkins "
        "Hospital, Massachusetts General Hospital, University College London Hospital, and Seoul "
        "National University Hospital."
    )

    doc.add_page_break()

    # --- Chapter 4: Results ---
    doc.add_heading("Chapter 4: Results", level=1)
    doc.add_paragraph(
        "Our meta-analysis revealed significant performance improvements across all examined "
        "healthcare domains. In diagnostic imaging, ML models achieved a pooled AUC-ROC of 0.947 "
        "(95% CI: 0.931-0.963) compared to 0.892 (95% CI: 0.871-0.913) for human experts alone. "
        "The improvement was most pronounced in retinal disease detection, where deep learning "
        "models identified diabetic retinopathy with 96.8% sensitivity."
    )
    doc.add_paragraph(
        "In drug discovery, ML-accelerated virtual screening reduced the average time from target "
        "identification to lead compound selection from 4.5 years to approximately 18 months. "
        "Notable examples include Insilico Medicine's identification of a novel DDR1 kinase "
        "inhibitor in just 46 days using generative adversarial networks."
    )
    doc.add_paragraph(
        "Clinical interview participants expressed cautious optimism about AI integration. "
        "Dr. Sarah Chen, a radiologist at Stanford, noted: 'AI has become an indispensable "
        "second reader in our workflow. It catches findings I might miss at the end of a long "
        "shift, but it also produces false positives that require clinical judgment to filter.'"
    )

    doc.add_page_break()

    # --- Chapter 5: Discussion ---
    doc.add_heading("Chapter 5: Discussion", level=1)
    doc.add_paragraph(
        "The results of this study corroborate the growing consensus that machine learning can "
        "substantially augment clinical decision-making, particularly in pattern recognition tasks. "
        "However, several critical challenges remain before widespread clinical deployment can be "
        "considered safe and equitable."
    )
    doc.add_paragraph(
        "First, the issue of algorithmic bias continues to pose significant risks. Studies by "
        "Obermeyer et al. (2019) demonstrated that a widely used healthcare algorithm exhibited "
        "racial bias, systematically underestimating the health needs of Black patients. Our "
        "review found that only 31% of the examined studies reported demographic breakdowns of "
        "their training data, and fewer than 15% conducted formal fairness audits."
    )
    doc.add_paragraph(
        "Second, the interpretability of complex ML models remains a barrier to clinical adoption. "
        "While attention visualization and gradient-based methods like Grad-CAM provide some "
        "insight into model reasoning, clinicians in our interviews expressed frustration with "
        "the 'black box' nature of many deployed systems."
    )

    doc.add_page_break()

    # --- Chapter 6: Conclusion ---
    doc.add_heading("Chapter 6: Conclusion", level=1)
    doc.add_paragraph(
        "This comprehensive review has demonstrated that machine learning holds transformative "
        "potential for healthcare, with documented improvements in diagnostic accuracy, drug "
        "discovery efficiency, and treatment personalization. The evidence from 147 studies and "
        "46 clinical interviews paints a picture of a field at an inflection point."
    )
    doc.add_paragraph(
        "We recommend a three-pronged approach for the responsible integration of ML in "
        "healthcare: (1) mandatory bias auditing and demographic reporting for all clinical AI "
        "systems, (2) development of interpretability standards that meet clinician needs, and "
        "(3) establishment of regulatory frameworks that balance innovation with patient safety."
    )
    doc.add_paragraph(
        "Future work should focus on longitudinal studies that track patient outcomes in "
        "ML-augmented clinical settings over periods of five years or more. Only through "
        "sustained observation can we truly understand the long-term impact of artificial "
        "intelligence on healthcare quality, equity, and cost-effectiveness."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
