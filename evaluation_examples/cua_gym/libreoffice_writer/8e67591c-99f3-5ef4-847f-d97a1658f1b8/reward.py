"""
Reward Script: Lab report creation in LibreOffice Writer
Task ID: writer_wf_025
Domain: libreoffice_writer
Scoring:
  C1: Title/Author/Date present (0.15)
  C2: Abstract italic ~100 words (0.15)
  C3: 7 Heading 1 sections (0.20)
  C4: Results table 5x4 with correct headers (0.20)
  C5: Double-spacing (0.15)
  C6: 2 reference entries (0.15)
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_025'


def persist_app_state(domain: str):
    """Try to save any unsaved LibreOffice state."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    paragraphs = doc.paragraphs
    if len(paragraphs) == 0:
        print("FAIL: Document has no paragraphs — appears blank")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title, Author, Date present (0.15 points)
    # Task says title='Determination of Vitamin C Content in Citrus Fruits', author='Lab Group 4', date present
    try:
        all_text = [p.text.strip() for p in paragraphs]
        has_title = any('vitamin c' in t.lower() and 'citrus' in t.lower() for t in all_text)
        has_author = any('lab group 4' in t.lower() for t in all_text)
        # Date: any paragraph with a date-like pattern in the first 5 paragraphs
        import re
        has_date = any(re.search(r'\d{4}|\d{1,2}[/\-\.]\d{1,2}', t) for t in all_text[:6])

        if has_title and has_author and has_date:
            print(f"PASS: Component 1 — Title, Author, Date all present (0.15 pts)")
            total_score += 0.15
        else:
            missing = []
            if not has_title:
                missing.append("title")
            if not has_author:
                missing.append("author 'Lab Group 4'")
            if not has_date:
                missing.append("date")
            print(f"FAIL: Component 1 — Missing: {', '.join(missing)}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Abstract section with italic text, ~100 words (0.15 points)
    try:
        # Find the abstract heading and the paragraph after it
        abstract_idx = None
        for i, p in enumerate(paragraphs):
            if p.style.name == 'Heading 1' and 'abstract' in p.text.lower():
                abstract_idx = i
                break

        if abstract_idx is not None:
            # Find the body paragraph(s) after the abstract heading
            abstract_body = []
            for p in paragraphs[abstract_idx + 1:]:
                if p.style.name == 'Heading 1':
                    break
                if p.text.strip():
                    abstract_body.append(p)

            if abstract_body:
                # Check italic on runs of the abstract body paragraph(s)
                first_body = abstract_body[0]
                italic_runs = [r for r in first_body.runs if r.font.italic]
                all_runs_with_text = [r for r in first_body.runs if r.text.strip()]
                is_italic = len(italic_runs) > 0 and len(italic_runs) >= len(all_runs_with_text) * 0.5

                # Check word count (~100 words, allow 60-150 range)
                full_text = ' '.join(p.text for p in abstract_body)
                word_count = len(full_text.split())
                reasonable_length = 60 <= word_count <= 200

                if is_italic and reasonable_length:
                    print(f"PASS: Component 2 — Abstract is italic, {word_count} words (0.15 pts)")
                    total_score += 0.15
                elif is_italic:
                    print(f"PARTIAL: Component 2 — Abstract italic but word count {word_count} outside range (0.075 pts)")
                    total_score += 0.075
                elif reasonable_length:
                    print(f"PARTIAL: Component 2 — Abstract length OK ({word_count} words) but not italic (0.075 pts)")
                    total_score += 0.075
                else:
                    print(f"FAIL: Component 2 — Abstract not italic and word count {word_count}")
            else:
                print("FAIL: Component 2 — Abstract heading found but no body text")
        else:
            print("FAIL: Component 2 — No 'Abstract' heading found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: 7 Heading 1 sections with correct names (0.20 points)
    try:
        expected_sections = ['abstract', 'introduction', 'materials and methods',
                             'results', 'discussion', 'conclusion', 'references']
        h1_paragraphs = [p for p in paragraphs if p.style.name == 'Heading 1']
        h1_texts = [p.text.strip().lower() for p in h1_paragraphs]

        matched = 0
        for expected in expected_sections:
            if any(expected in h for h in h1_texts):
                matched += 1

        if matched == 7:
            print(f"PASS: Component 3 — All 7 Heading 1 sections found (0.20 pts)")
            total_score += 0.20
        elif matched >= 5:
            partial = round(0.20 * (matched / 7), 3)
            print(f"PARTIAL: Component 3 — {matched}/7 sections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {matched}/7 expected sections. Found headings: {h1_texts}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Results table with 4 columns, 5 rows (header + 4 fruits) (0.20 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)

            # Check dimensions
            correct_dims = num_rows >= 5 and num_cols >= 4

            # Check header row has expected column names
            headers = [table.cell(0, c).text.strip().lower() for c in range(min(num_cols, 4))]
            has_fruit_col = any('fruit' in h for h in headers)
            has_weight_col = any('weight' in h or 'sample' in h for h in headers)
            has_titration_col = any('titration' in h or 'volume' in h for h in headers)
            has_vitc_col = any('vitamin' in h or 'mg' in h for h in headers)
            headers_ok = has_fruit_col and has_weight_col and has_titration_col and has_vitc_col

            # Check data rows have 4 fruit entries
            fruit_names = []
            for r in range(1, min(num_rows, 6)):
                cell_text = table.cell(r, 0).text.strip().lower()
                if cell_text:
                    fruit_names.append(cell_text)
            has_4_fruits = len(fruit_names) >= 4

            if correct_dims and headers_ok and has_4_fruits:
                print(f"PASS: Component 4 — Table {num_rows}x{num_cols}, headers correct, {len(fruit_names)} fruits (0.20 pts)")
                total_score += 0.20
            elif correct_dims and (headers_ok or has_4_fruits):
                print(f"PARTIAL: Component 4 — Table dimensions OK, partial content match (0.10 pts)")
                total_score += 0.10
            elif correct_dims:
                print(f"PARTIAL: Component 4 — Table dimensions OK but headers/content wrong (0.05 pts)")
                total_score += 0.05
            else:
                print(f"FAIL: Component 4 — Table dims {num_rows}x{num_cols}, expected >=5x4")
        else:
            print("FAIL: Component 4 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Double-spacing (line_spacing=2.0) on body paragraphs (0.15 points)
    try:
        body_paras = [p for p in paragraphs if p.text.strip() and p.style.name != 'Heading 1']
        if body_paras:
            double_spaced_count = 0
            for p in body_paras:
                ls = p.paragraph_format.line_spacing
                if ls is not None and abs(float(ls) - 2.0) < 0.1:
                    double_spaced_count += 1

            ratio = double_spaced_count / len(body_paras) if body_paras else 0
            if ratio >= 0.7:
                print(f"PASS: Component 5 — {double_spaced_count}/{len(body_paras)} body paragraphs double-spaced (0.15 pts)")
                total_score += 0.15
            elif ratio >= 0.4:
                print(f"PARTIAL: Component 5 — {double_spaced_count}/{len(body_paras)} body paragraphs double-spaced (0.075 pts)")
                total_score += 0.075
            else:
                print(f"FAIL: Component 5 — Only {double_spaced_count}/{len(body_paras)} body paragraphs double-spaced")
        else:
            print("FAIL: Component 5 — No body paragraphs found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: References section with 2 entries (0.15 points)
    try:
        ref_idx = None
        for i, p in enumerate(paragraphs):
            if p.style.name == 'Heading 1' and 'reference' in p.text.lower():
                ref_idx = i
                break

        if ref_idx is not None:
            ref_entries = [p for p in paragraphs[ref_idx + 1:] if p.text.strip()]
            if len(ref_entries) >= 2:
                print(f"PASS: Component 6 — References section with {len(ref_entries)} entries (0.15 pts)")
                total_score += 0.15
            elif len(ref_entries) == 1:
                print(f"PARTIAL: Component 6 — Only 1 reference entry (0.075 pts)")
                total_score += 0.075
            else:
                print(f"FAIL: Component 6 — References heading found but no entries")
        else:
            print("FAIL: Component 6 — No 'References' heading found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
