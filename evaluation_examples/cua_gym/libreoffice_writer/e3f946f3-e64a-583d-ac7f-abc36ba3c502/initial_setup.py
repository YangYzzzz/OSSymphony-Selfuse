"""
Initial Setup: Apply superscript/subscript formatting to ordinals and chemical formulas
Task ID: writer_rd_092
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_092'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Annual Environmental Chemistry Report', level=0)

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This report presents the findings from our 3rd annual survey of water quality '
        'and atmospheric composition across the Pacific Northwest region. The study was '
        'conducted between March 1st and September 30th, 2025, covering 12 monitoring '
        'stations. This marks the 2nd consecutive year where we have expanded our '
        'sampling methodology to include deep-water analysis.'
    )

    # --- Section 2: Water Quality Analysis ---
    doc.add_heading('2. Water Quality Analysis', level=1)
    doc.add_paragraph(
        'Water samples were collected from 8 different sites. The 1st set of samples '
        'was taken from Lake Meridian, where dissolved H2O purity levels exceeded 99.7%. '
        'We detected trace amounts of NaCl at concentrations of 0.03 mol/L, consistent '
        'with natural mineral leaching from surrounding geological formations.'
    )
    doc.add_paragraph(
        'The 4th sampling station at Clearwater Creek showed elevated levels of CO2 '
        'dissolved in the water column, measuring 2.8 mg/L. This was the 21st consecutive '
        'month of rising CO2 concentrations at this location. Additionally, methane (CH4) '
        'was detected at 0.15 mg/L in the sediment layer, which represents the 2nd highest '
        'reading in our 10-year dataset.'
    )

    # --- Section 3: Atmospheric Composition ---
    doc.add_heading('3. Atmospheric Composition', level=1)
    doc.add_paragraph(
        'Atmospheric monitoring began on the 3rd of March at an elevation of 1,450 meters. '
        'The 5th monitoring tower recorded CO2 levels of 421 ppm, while CH4 concentrations '
        'were measured at 1,892 ppb. These findings place our region in the 7th percentile '
        'nationally for greenhouse gas concentrations.'
    )
    doc.add_paragraph(
        'Sulfur dioxide (SO2) emissions from the nearby industrial zone averaged 12.3 ppb, '
        'marking the 3rd year of decline following the implementation of new filtration '
        'systems. Nitrous oxide (N2O) levels remained stable at 332 ppb throughout the '
        'monitoring period.'
    )

    # --- Section 4: Chemical Analysis Summary ---
    doc.add_heading('4. Chemical Analysis Summary', level=1)

    # Add a table with chemical data
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Compound', 'Location', 'Concentration', 'Status']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    data = [
        ['H2O', 'Lake Meridian', '99.7% purity', 'Normal'],
        ['CO2', 'Clearwater Creek', '2.8 mg/L', 'Elevated'],
        ['NaCl', 'Lake Meridian', '0.03 mol/L', 'Normal'],
        ['CH4', 'Clearwater Creek', '0.15 mg/L', 'Elevated'],
        ['SO2', 'Industrial Zone', '12.3 ppb', 'Declining'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')  # spacer

    # --- Section 5: Conclusions ---
    doc.add_heading('5. Conclusions', level=1)
    doc.add_paragraph(
        'This study represents the 3rd comprehensive assessment of environmental chemistry '
        'in our region. For the 1st time, we observed a statistically significant correlation '
        'between dissolved CO2 and N2O levels in freshwater systems. The 9th recommendation '
        'from our advisory board, regarding enhanced H2O filtration protocols, has been '
        'implemented at all monitoring stations since the 4th quarter of 2024.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
