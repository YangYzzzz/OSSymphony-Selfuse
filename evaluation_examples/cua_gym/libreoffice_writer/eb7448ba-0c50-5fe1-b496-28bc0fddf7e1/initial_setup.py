"""
Initial Setup: Safety Manual document - pre-task state (no shapes)
Task ID: writer_obj_073
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_073'
OUTPUT = f'{WORKDIR}/safety_manual.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading('Workplace Safety Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')

    # Section 1
    doc.add_heading('1. General Safety Guidelines', level=1)
    p1 = doc.add_paragraph(
        'All employees must adhere to the following safety guidelines at all times. '
        'Failure to comply may result in injury, disciplinary action, or termination. '
        'Safety is a shared responsibility and every team member plays a critical role '
        'in maintaining a safe working environment.'
    )

    doc.add_paragraph(
        'Personal Protective Equipment (PPE) must be worn in designated areas. '
        'Hard hats are required on the construction floor. Safety goggles must be '
        'used when operating machinery or handling chemicals. High-visibility vests '
        'are mandatory in vehicle access zones.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Emergency exits must remain unobstructed at all times. Exit routes are marked '
        'with illuminated green signs and must not be blocked by equipment, materials, '
        'or furniture.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'All incidents, near-misses, and hazards must be reported immediately to the '
        'safety officer or direct supervisor using the Incident Report Form (IRF-001).',
        style='List Bullet'
    )

    # Section 2
    doc.add_heading('2. Hazard Identification and Risk Assessment', level=1)
    doc.add_paragraph(
        'Before beginning any task, employees must conduct a Job Hazard Analysis (JHA). '
        'This involves identifying potential hazards, evaluating the risk level, and '
        'implementing appropriate control measures to reduce the likelihood of accidents.'
    )
    doc.add_paragraph(
        'Chemical Hazards: All chemical substances must be stored in clearly labeled '
        'containers. Safety Data Sheets (SDS) must be accessible for all chemicals in use.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Physical Hazards: Slipping, tripping, and falling are among the most common '
        'workplace accidents. Keep walkways clean and dry. Report spills immediately.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Electrical Hazards: Only qualified electricians are permitted to perform '
        'electrical work. Do not use damaged cords or outlets. Tag and lock out '
        'all electrical equipment before maintenance.',
        style='List Bullet'
    )

    # Section 3
    doc.add_heading('3. Emergency Procedures', level=1)
    doc.add_paragraph(
        'In the event of an emergency, follow the evacuation procedures outlined below. '
        'The assembly point is located in Parking Lot B, adjacent to the main entrance. '
        'Do not use elevators during an emergency evacuation.'
    )

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Emergency Type'
    hdr[1].text = 'Primary Action'
    hdr[2].text = 'Contact'
    for cell in hdr:
        run = cell.paragraphs[0].runs[0]
        run.bold = True

    rows_data = [
        ('Fire', 'Activate alarm, evacuate immediately', 'Fire Safety: ext. 119'),
        ('Chemical Spill', 'Evacuate area, call HAZMAT team', 'HAZMAT Hotline: ext. 221'),
        ('Medical Emergency', 'Call first aid, do not move victim', 'First Aid: ext. 112'),
    ]
    for i, (t, a, c) in enumerate(rows_data, 1):
        row = table.rows[i].cells
        row[0].text = t
        row[1].text = a
        row[2].text = c

    doc.add_paragraph('')

    # Section 4
    doc.add_heading('4. Safety Training Requirements', level=1)
    doc.add_paragraph(
        'All new employees must complete the mandatory Safety Induction Program (SIP) '
        'within the first week of employment. Refresher training is conducted annually. '
        'Specialized training is required for roles involving heavy machinery, '
        'hazardous materials, or elevated work platforms.'
    )

    doc.add_paragraph(
        'Induction Training: 8-hour program covering facility layout, emergency procedures, '
        'PPE usage, and reporting protocols. Completion certificate required before '
        'accessing the production floor.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Annual Refresher: 4-hour update session covering regulatory changes, incident '
        'reviews, and updated safety procedures. All employees required annually by '
        'December 31st.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Specialized Courses: Forklift Operation (FO-301), Working at Heights (WAH-201), '
        'Confined Space Entry (CSE-401), and Chemical Handling (CH-101).',
        style='List Bullet'
    )

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = 'Safety Manual v2.4 | Approved: 2025-01-15 | Next Review: 2026-01-15'
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fp.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
