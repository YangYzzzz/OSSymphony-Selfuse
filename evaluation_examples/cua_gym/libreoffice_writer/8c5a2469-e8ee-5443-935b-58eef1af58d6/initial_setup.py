"""
Initial Setup: photo_album.docx — photo album document with image, no border/shadow
Task ID: writer_obj_033
Domain: libreoffice_writer
"""

import os
import io
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image, ImageDraw

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_033'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/photo_album.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_photo_png(width_px=300, height_px=225, color=(120, 160, 200)):
    """Create a realistic-looking photo PNG in memory."""
    img = Image.new('RGB', (width_px, height_px), color=color)
    draw = ImageDraw.Draw(img)
    # Add some visual interest — simple landscape scene
    # Sky gradient area
    for y in range(height_px // 2):
        r = int(100 + y * 0.4)
        g = int(150 + y * 0.3)
        b = int(220 - y * 0.3)
        draw.line([(0, y), (width_px, y)], fill=(r, g, b))
    # Ground
    for y in range(height_px // 2, height_px):
        r = int(60 + (y - height_px // 2) * 0.5)
        g = int(100 + (y - height_px // 2) * 0.3)
        b = int(40)
        draw.line([(0, y), (width_px, y)], fill=(r, g, b))
    # Simple tree silhouette
    draw.ellipse([100, 50, 200, 130], fill=(30, 80, 30))
    draw.rectangle([140, 120, 160, 160], fill=(80, 50, 20))
    # Simple sun
    draw.ellipse([220, 20, 260, 60], fill=(255, 230, 50))
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Page 1 heading
    heading = doc.add_heading('Summer Vacation 2024', level=1)
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading.paragraph_format.space_after = Pt(12)

    # Introductory text
    intro = doc.add_paragraph(
        'A wonderful collection of memories from our trip to the coast. '
        'We visited Maplewood Beach on the first day, where the weather was '
        'sunny and warm. The kids had a great time exploring the tide pools '
        'and building sand castles near the shoreline.'
    )
    intro.paragraph_format.space_after = Pt(6)

    # Caption before image
    caption_before = doc.add_paragraph('Morning at Maplewood Beach — Day 1')
    caption_before.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption_before.paragraph_format.space_before = Pt(6)
    caption_before.paragraph_format.space_after = Pt(4)
    run_caption = caption_before.runs[0]
    run_caption.font.italic = True
    run_caption.font.size = Pt(10)

    # Add photograph: 8cm x 6cm, NO border, NO shadow
    photo_para = doc.add_paragraph()
    photo_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    photo_para.paragraph_format.space_before = Pt(4)
    photo_para.paragraph_format.space_after = Pt(4)
    photo_run = photo_para.add_run()
    photo_buf = make_photo_png(480, 360)  # 4:3 ratio, will be scaled to 8cm x 6cm
    photo_run.add_picture(photo_buf, width=Cm(8), height=Cm(6))

    # Caption after image
    caption_after = doc.add_paragraph('Photo 1: Scenic view of the coastline at dawn.')
    caption_after.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption_after.paragraph_format.space_before = Pt(4)
    caption_after.paragraph_format.space_after = Pt(12)
    run_ca = caption_after.runs[0]
    run_ca.font.size = Pt(9)
    run_ca.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # Additional paragraph on page 1
    p2 = doc.add_paragraph(
        'The afternoon was spent at the local farmers market, browsing fresh '
        'produce and handmade crafts. We picked up some locally sourced honey '
        'and a few lavender sachets to bring back home. The evening ended with '
        'a beautiful sunset dinner by the water.'
    )
    p2.paragraph_format.space_after = Pt(8)

    # Page break before page 2
    doc.add_page_break()

    # Page 2
    heading2 = doc.add_heading('Day 2 — Mountain Hike', level=1)
    heading2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading2.paragraph_format.space_after = Pt(12)

    p3 = doc.add_paragraph(
        'Our second day took us inland to Ridgecrest Trail. The hike began '
        'early in the morning, before the summer heat became too intense. '
        'Along the path we spotted deer and several species of native birds. '
        'The view from the summit overlooked three valleys and the distant '
        'Pacific coastline. It was absolutely breathtaking.'
    )
    p3.paragraph_format.space_after = Pt(8)

    p4 = doc.add_paragraph(
        'We packed sandwiches and lemonade for lunch, which we enjoyed at a '
        'picnic table near Clearwater Creek. The children waded in the shallow '
        'water while the adults relaxed in the shade. By late afternoon we '
        'made our way back to the trailhead, tired but happy.'
    )
    p4.paragraph_format.space_after = Pt(8)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
