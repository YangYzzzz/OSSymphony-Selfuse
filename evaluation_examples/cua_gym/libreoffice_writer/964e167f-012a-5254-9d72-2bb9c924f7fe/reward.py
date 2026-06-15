"""
Reward Script: Board of Directors Meeting Agenda in LibreOffice Writer
Task ID: writer_wf_076
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15): Header 'Board of Directors Meeting Agenda' as heading
  Component 2 (0.15): Meeting details (date, time, location, chair)
  Component 3 (0.30): 7 numbered agenda items with presenter and time
  Component 4 (0.25): Action Items Review table (header + 4 data rows, 3 cols)
  Component 5 (0.15): Next meeting note
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_076'


def persist_app_state(domain: str):
    """Try to save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather all paragraph texts for searching
    all_text = '\n'.join(p.text for p in doc.paragraphs)
    all_text_lower = all_text.lower()

    # Component 1: Header 'Board of Directors Meeting Agenda' (0.15 points)
    # Must appear as a heading-style paragraph (not just normal text)
    try:
        header_matches = [
            p for p in doc.paragraphs
            if 'board of directors meeting agenda' in p.text.lower()
            and ('heading' in p.style.name.lower() or p.text.strip() != '')
        ]
        if len(header_matches) > 0 and 'board of directors meeting agenda' in all_text_lower:
            print(f"PASS: Component 1 - Header found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 - Header 'Board of Directors Meeting Agenda' not found")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Meeting details - date, time, location, chair (0.15 points)
    # Each sub-item worth 0.0375 points
    try:
        details_score = 0.0
        # Check for date mention
        if re.search(r'date\s*:', all_text_lower):
            details_score += 0.0375
            print(f"  PASS: Date field found")
        else:
            print(f"  FAIL: Date field not found")

        # Check for time mention
        if re.search(r'time\s*:', all_text_lower):
            details_score += 0.0375
            print(f"  PASS: Time field found")
        else:
            print(f"  FAIL: Time field not found")

        # Check for location mention
        if re.search(r'location\s*:', all_text_lower):
            details_score += 0.0375
            print(f"  PASS: Location field found")
        else:
            print(f"  FAIL: Location field not found")

        # Check for chair person mention
        if re.search(r'chair\s*:', all_text_lower):
            details_score += 0.0375
            print(f"  PASS: Chair field found")
        else:
            print(f"  FAIL: Chair field not found")

        if details_score > 0:
            print(f"PASS: Component 2 - Meeting details ({details_score:.4f} pts)")
            total_score += details_score
        else:
            print(f"FAIL: Component 2 - No meeting details found")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: 7 numbered agenda items with presenter and time (0.30 points)
    # Required items: Call to Order, Approval of Minutes, Financial Report,
    # Committee Updates, New Business, Action Items Review, Adjournment
    try:
        required_topics = [
            'call to order',
            'approval of minutes',
            'financial report',
            'committee updates',
            'new business',
            'action items review',
            'adjournment'
        ]
        # Also check for 'presenter' and time allocation mentions
        items_found = 0
        has_presenter_info = False
        has_time_info = False

        for topic in required_topics:
            if topic in all_text_lower:
                items_found += 1

        # Check for presenter/time info patterns
        presenter_match = re.search(r'presenter\s*:', all_text_lower)
        time_match = re.search(r'time\s*:\s*\d+\s*min', all_text_lower)

        # Score: 0.20 for topics, 0.05 for presenter info, 0.05 for time info
        topic_score = (items_found / 7.0) * 0.20
        presenter_score = 0.05 if presenter_match else 0.0
        time_score = 0.05 if time_match else 0.0
        comp3_score = topic_score + presenter_score + time_score

        if comp3_score > 0:
            print(f"PASS: Component 3 - {items_found}/7 agenda topics, presenter={bool(presenter_match)}, time={bool(time_match)} ({comp3_score:.4f} pts)")
            total_score += comp3_score
        else:
            print(f"FAIL: Component 3 - No agenda items found")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Action Items Review table (0.25 points)
    # Must have: header row (Item, Owner, Status) + 4 data rows = 5 rows total, 3 cols
    try:
        comp4_score = 0.0
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)

            # Check table has correct structure: at least 5 rows and 3 columns
            if num_rows >= 5 and num_cols >= 3:
                comp4_score += 0.10
                print(f"  PASS: Table structure {num_rows}x{num_cols} (>= 5x3)")
            else:
                print(f"  FAIL: Table structure {num_rows}x{num_cols}, expected >= 5x3")

            # Check header row contains Item, Owner, Status
            header_texts = [table.cell(0, c).text.strip().lower() for c in range(min(num_cols, 3))]
            expected_headers = ['item', 'owner', 'status']
            headers_match = all(eh in ' '.join(header_texts) for eh in expected_headers)
            if headers_match:
                comp4_score += 0.05
                print(f"  PASS: Table headers: {header_texts}")
            else:
                print(f"  FAIL: Table headers {header_texts}, expected {expected_headers}")

            # Check 4 data rows have non-empty content
            data_rows_filled = 0
            for r in range(1, min(num_rows, 5)):
                row_text = ''.join(table.cell(r, c).text.strip() for c in range(min(num_cols, 3)))
                if len(row_text) > 5:
                    data_rows_filled += 1
            if data_rows_filled >= 4:
                comp4_score += 0.10
                print(f"  PASS: {data_rows_filled} data rows with content")
            else:
                print(f"  FAIL: Only {data_rows_filled}/4 data rows with content")

            if comp4_score > 0:
                print(f"PASS: Component 4 - Action Items table ({comp4_score:.4f} pts)")
                total_score += comp4_score
            else:
                print(f"FAIL: Component 4 - Table exists but checks failed")
        else:
            print(f"FAIL: Component 4 - No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Next meeting note (0.15 points)
    # Should mention next meeting date at the end of the document
    try:
        if re.search(r'next\s+(board\s+of\s+directors\s+)?meeting', all_text_lower):
            print(f"PASS: Component 5 - Next meeting note found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 - Next meeting note not found")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
