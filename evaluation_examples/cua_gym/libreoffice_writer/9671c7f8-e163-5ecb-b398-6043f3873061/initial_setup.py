"""
Initial Setup: Create ranking.docx with numbered list using standard "1." format
Task ID: writer_list_057
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_list_057'
OUTPUT = f'{WORKDIR}/Desktop/ranking.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Add a title/heading for context
    heading = doc.add_paragraph('Strategic Priorities', style='Heading 1')

    # Add five numbered list items using standard "1." format
    items = [
        'Customer satisfaction improvement',
        'Revenue growth targets',
        'Employee development programs',
        'Technology infrastructure upgrade',
        'Supply chain optimization',
    ]

    for item in items:
        doc.add_paragraph(item, style='List Number')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
