"""
Reward Script: Create wine bottle labels with custom size and merge fields
Task ID: writer_mt_036
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table exists as label grid (rows >= 2, cols >= 1)
  Component 2 (0.25): Merge fields WineName, Vintage, Region present in label cells
  Component 3 (0.25): Font formatting — WineName 14pt bold, Vintage 12pt, Region 10pt italic
  Component 4 (0.25): Label dimensions — row height ~2 inches, column width ~3.5 inches
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_036'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            import time
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists as label grid (0.25 points)
    # The label document should have at least one table with multiple rows/cols forming a label grid.
    # The initial document has 0 tables, so this differentiates initial from golden.
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows >= 2 and num_cols >= 1:
                print(f"PASS: Component 1 — Label table exists with {num_rows} rows x {num_cols} cols (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 — Table too small: {num_rows} rows x {num_cols} cols, need >= 2 rows and >= 1 col")
        else:
            print(f"FAIL: Component 1 — No tables found in document (found {len(doc.tables)})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Gate: need at least one table to check remaining components
    if len(doc.tables) < 1:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    table = doc.tables[0]

    # Component 2: Merge fields WineName, Vintage, Region in label cells (0.25 points)
    # Each label cell should contain the three merge field placeholders.
    # Check the first non-empty cell for all three fields.
    try:
        fields_found = {"WineName": False, "Vintage": False, "Region": False}
        # Check first cell
        cell = table.rows[0].cells[0]
        cell_text = cell.text
        for field in fields_found:
            if field in cell_text or f"<{field}>" in cell_text or f"{{{field}}}" in cell_text:
                fields_found[field] = True

        all_fields = all(fields_found.values())
        if all_fields:
            print(f"PASS: Component 2 — All merge fields found: WineName, Vintage, Region (0.25 pts)")
            total_score += 0.25
        else:
            missing = [k for k, v in fields_found.items() if not v]
            print(f"FAIL: Component 2 — Missing merge fields: {missing}. Cell text: {repr(cell_text[:200])}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Font formatting (0.25 points)
    # WineName line: 14pt bold; Vintage line: 12pt; Region line: 10pt italic
    # Check the first cell's paragraphs for correct formatting.
    try:
        cell = table.rows[0].cells[0]
        paras = cell.paragraphs
        formatting_score = 0.0
        checks_detail = []

        # We need at least 3 paragraphs (one per field)
        if len(paras) >= 3:
            passed_count = 0  # count of formatting checks that pass

            # Check WineName paragraph (first): should be 14pt bold
            wine_para = paras[0]
            for run in wine_para.runs:
                if "WineName" in run.text or "Wine" in run.text:
                    is_bold = run.font.bold is True
                    is_14pt = run.font.size is not None and abs(run.font.size.pt - 14.0) < 0.5
                    if is_bold and is_14pt:
                        passed_count += 1
                        checks_detail.append(f"WineName: bold={is_bold}, size={run.font.size.pt}pt OK")
                    else:
                        checks_detail.append(f"WineName: bold={run.font.bold}, size={run.font.size.pt if run.font.size else None}pt")
                    break

            # Check Vintage paragraph (second): should be 12pt
            vintage_para = paras[1]
            for run in vintage_para.runs:
                if "Vintage" in run.text:
                    is_12pt = run.font.size is not None and abs(run.font.size.pt - 12.0) < 0.5
                    if is_12pt:
                        passed_count += 1
                        checks_detail.append(f"Vintage: size={run.font.size.pt}pt OK")
                    else:
                        checks_detail.append(f"Vintage: size={run.font.size.pt if run.font.size else None}pt")
                    break

            # Check Region paragraph (third): should be 10pt italic
            region_para = paras[2]
            for run in region_para.runs:
                if "Region" in run.text:
                    is_italic = run.font.italic is True
                    is_10pt = run.font.size is not None and abs(run.font.size.pt - 10.0) < 0.5
                    if is_italic and is_10pt:
                        passed_count += 1
                        checks_detail.append(f"Region: italic={is_italic}, size={run.font.size.pt}pt OK")
                    else:
                        checks_detail.append(f"Region: italic={run.font.italic}, size={run.font.size.pt if run.font.size else None}pt")
                    break

            if passed_count == 3:
                formatting_score = 0.25
            elif passed_count == 2:
                formatting_score = 0.15
            elif passed_count == 1:
                formatting_score = 0.08

            if formatting_score > 0:
                print(f"PASS: Component 3 — Font formatting verified ({formatting_score} pts): {'; '.join(checks_detail)}")
                total_score += formatting_score
            else:
                print(f"FAIL: Component 3 — Font formatting incorrect: {'; '.join(checks_detail)}")
        else:
            print(f"FAIL: Component 3 — First cell has {len(paras)} paragraphs, need >= 3 for WineName/Vintage/Region")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Label dimensions (0.25 points)
    # Row height should be ~2 inches (2880 twips), column width ~3.5 inches (5040 twips)
    # Allow some tolerance since page layout may adjust widths.
    try:
        dim_score = 0.0

        # Check row height (should be ~2 inches = 2880 twips)
        row_heights_ok = 0
        row_heights_total = 0
        for row in table.rows:
            tr = row._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is not None:
                trHeight = trPr.find(qn('w:trHeight'))
                if trHeight is not None:
                    h_val = trHeight.get(qn('w:val'))
                    if h_val:
                        h_inches = int(h_val) / 1440.0
                        row_heights_total += 1
                        # Allow tolerance: 1.5 to 2.5 inches
                        if 1.5 <= h_inches <= 2.5:
                            row_heights_ok += 1

        # Check column width (should be ~3.5 inches = 5040 twips)
        col_widths_ok = 0
        col_widths_total = 0
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    w_val = tcW.get(qn('w:w'))
                    if w_val:
                        w_inches = int(w_val) / 1440.0
                        col_widths_total += 1
                        # Allow tolerance: 3.0 to 4.5 inches (labels can be slightly wider)
                        if 3.0 <= w_inches <= 4.5:
                            col_widths_ok += 1

        height_ok = row_heights_total > 0 and row_heights_ok == row_heights_total
        width_ok = col_widths_total > 0 and col_widths_ok == col_widths_total

        if height_ok and width_ok:
            dim_score = 0.25
            print(f"PASS: Component 4 — Label dimensions correct: rows ~2\" tall, cols ~3.5\" wide (0.25 pts)")
        elif height_ok or width_ok:
            dim_score = 0.12
            detail = "height OK" if height_ok else "width OK"
            print(f"PARTIAL: Component 4 — {detail} ({dim_score} pts)")
        else:
            print(f"FAIL: Component 4 — Dimensions wrong: row heights OK={row_heights_ok}/{row_heights_total}, col widths OK={col_widths_ok}/{col_widths_total}")

        if dim_score > 0:
            total_score += dim_score
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
