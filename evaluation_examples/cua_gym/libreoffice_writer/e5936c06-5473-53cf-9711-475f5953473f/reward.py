"""
Reward Script: Employee Evaluation Form in Writer
Task ID: writer_pd_011
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Employee info table — 4 rows x 2 cols with labels
  Component 2 (0.25): Rating table — header + 8 criteria rows x 2 cols
  Component 3 (0.30): Form fields — 4 text inputs, 8 dropdown fields, 1 comments text field
  Component 4 (0.20): Document protection — forms protection enabled
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_011'


def persist_app_state(domain: str):
    """Try to save any unsaved changes in LibreOffice."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body = doc.element.body

    # =========================================================================
    # Component 1: Employee info table (0.25 points)
    # Must have a table with 4 rows x 2 cols containing Name, Department,
    # Manager, Date labels. This table does NOT exist in initial_env.
    # =========================================================================
    try:
        expected_labels = {"name", "department", "manager", "date"}
        info_labels_matched = 0

        for table in doc.tables:
            if len(table.rows) >= 4 and len(table.columns) >= 2:
                # Check if the first column contains the expected labels
                labels_found = set()
                for row in table.rows:
                    cell_text = row.cells[0].text.strip().lower()
                    for label in expected_labels:
                        if label in cell_text:
                            labels_found.add(label)
                if labels_found == expected_labels:
                    info_labels_matched = len(labels_found)
                    break

        if info_labels_matched >= 4:
            print(f"PASS: Component 1 — Employee info table found with all 4 labels (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Employee info table not found or missing labels")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # =========================================================================
    # Component 2: Rating table (0.25 points)
    # Must have a table with header row + 8 performance criteria rows (9 total),
    # at least 2 columns. This table does NOT exist in initial_env.
    # =========================================================================
    try:
        expected_criteria = [
            "job knowledge",
            "quality of work",
            "productivity",
            "communication",
            "teamwork",
            "initiative",
            "attendance",
            "problem solving",
        ]
        criteria_matched = 0

        for table in doc.tables:
            if len(table.rows) >= 9 and len(table.columns) >= 2:
                # Check if the first column contains performance criteria
                criteria_found = 0
                for row in table.rows:
                    cell_text = row.cells[0].text.strip().lower()
                    for criterion in expected_criteria:
                        if criterion in cell_text:
                            criteria_found += 1
                            break
                if criteria_found >= 7:  # Allow slight flexibility (7 of 8)
                    criteria_matched = criteria_found
                    break

        if criteria_matched >= 7:
            print(f"PASS: Component 2 — Rating table found with {criteria_found} criteria (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — Rating table not found or insufficient criteria (found {criteria_found if 'criteria_found' in dir() else 0})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =========================================================================
    # Component 3: Form fields (0.30 points)
    # Must have legacy form fields:
    #   - At least 4 text input fields for employee info (0.10)
    #   - At least 8 dropdown fields with 1-5 values for ratings (0.10)
    #   - At least 1 text input field for comments (0.10)
    # None of these exist in initial_env (0 form fields).
    # =========================================================================
    try:
        ff_datas = body.findall('.//w:ffData', ns)
        text_fields = []
        dropdown_fields = []

        for ff in ff_datas:
            text_input = ff.find('w:textInput', ns)
            dd_list = ff.find('w:ddList', ns)
            name_el = ff.find('w:name', ns)
            name = name_el.get(qn('w:val')) if name_el is not None else 'unnamed'

            if text_input is not None:
                text_fields.append(name)
            elif dd_list is not None:
                # Verify dropdown has 1-5 entries
                entries = dd_list.findall('w:listEntry', ns)
                vals = [e.get(qn('w:val')) for e in entries]
                dropdown_fields.append({'name': name, 'values': vals})

        # Sub-component 3a: Text input fields for employee info (0.10)
        info_text_count = sum(1 for n in text_fields if 'info' in n.lower() or 'name' in n.lower() or 'department' in n.lower() or 'manager' in n.lower() or 'date' in n.lower())
        # If names don't have 'info' prefix, just count text fields >= 4 (excluding comments-like)
        if info_text_count < 4:
            # Fallback: count all text fields minus 1 (for comments)
            info_text_count = max(0, len(text_fields) - 1) if len(text_fields) > 1 else len(text_fields)

        if info_text_count >= 4:
            print(f"PASS: Component 3a — Found {info_text_count} text input fields for employee info (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3a — Expected >= 4 text input fields for info, found {info_text_count}")

        # Sub-component 3b: Dropdown fields with 1-5 for ratings (0.10)
        valid_dropdowns = 0
        for dd in dropdown_fields:
            # Check that dropdown has entries containing 1-5
            dd_vals = dd['values']
            if len(dd_vals) >= 5:
                has_scale = all(str(v) in [str(x) for x in range(1, 6)] for v in dd_vals[:5])
                if has_scale:
                    valid_dropdowns += 1

        if valid_dropdowns >= 8:
            print(f"PASS: Component 3b — Found {valid_dropdowns} dropdown fields with 1-5 scale (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3b — Expected >= 8 dropdown fields with 1-5 scale, found {valid_dropdowns}")

        # Sub-component 3c: Comments text field (0.10)
        comments_field = any('comment' in n.lower() for n in text_fields)
        # Fallback: if there's at least one more text field beyond the 4 info ones
        if not comments_field and len(text_fields) >= 5:
            comments_field = len(text_fields) >= 5  # derived from count check

        if comments_field:
            print(f"PASS: Component 3c — Found comments text field (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3c — No comments text field found")

    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # =========================================================================
    # Component 4: Document protection for forms (0.20 points)
    # Must have documentProtection with edit=forms and enforcement=1.
    # Initial_env has NO protection.
    # =========================================================================
    try:
        settings_el = doc.settings.element
        prot = settings_el.find(qn('w:documentProtection'))

        if prot is not None:
            edit_val = prot.get(qn('w:edit'))
            enforce_val = prot.get(qn('w:enforcement'))

            # Check that protection is for forms and enforced
            if edit_val == 'forms' and enforce_val in ('1', 'true'):
                print(f"PASS: Component 4 — Document protection enabled (edit=forms, enforcement={enforce_val}) (0.20 pts)")
                total_score += 0.20
            elif edit_val == 'forms':
                print(f"FAIL: Component 4 — Protection type is 'forms' but enforcement={enforce_val}")
            else:
                print(f"FAIL: Component 4 — Protection type is '{edit_val}', expected 'forms'")
        else:
            print(f"FAIL: Component 4 — No document protection found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    # Also check for Eval_Form.docx as mentioned in task instruction
    file_path = f'{WORKDIR}/Eval_Form.docx'
    if not os.path.exists(file_path):
        print(f"File not found: {WORKDIR}/{TASK_ID}.docx or {WORKDIR}/Eval_Form.docx")
        print("REWARD: 0.0")
    else:
        verify_task(file_path)
else:
    verify_task(file_path)
