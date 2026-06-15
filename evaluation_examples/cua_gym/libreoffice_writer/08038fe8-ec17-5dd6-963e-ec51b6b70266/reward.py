"""
Reward Script: Certificate of Achievement document
Task ID: writer_wf_067
Domain: libreoffice_writer
Scoring:
  C1: Title text, centered, bold, ~26pt (0.20)
  C2: "This is to certify that" + name blank line ~20pt (0.15)
  C3: "has successfully completed" + program name (0.15)
  C4: Date and certificate number fields (0.15)
  C5: Signature table with Program Director and Dean (0.15)
  C6: Institution name at bottom (0.10)
  C7: Page border present (0.10)
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_067'


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice state."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print("PERSIST: ctrl+s sent for {}".format(domain))
        except Exception as e:
            print("PERSIST_WARN: save hook failed: {}".format(e))


def verify_task(file_path):
    """
    Verify certificate document creation with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file {}: {}".format(file_path, e))
        print("REWARD: 0.0")
        return 0.0

    all_text = [p.text.strip() for p in doc.paragraphs]
    all_text_lower = [t.lower() for t in all_text]

    # Component 1: Title "CERTIFICATE OF ACHIEVEMENT" centered, bold, ~26pt (0.20 points)
    try:
        title_found = False
        for p in doc.paragraphs:
            if 'certificate of achievement' in p.text.lower():
                is_centered = p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                has_bold = False
                has_large_size = False
                for r in p.runs:
                    if r.font.bold:
                        has_bold = True
                    if r.font.size and r.font.size.pt >= 24:
                        has_large_size = True
                if is_centered and has_bold and has_large_size:
                    print("PASS: Component 1 — Title 'CERTIFICATE OF ACHIEVEMENT' centered, bold, >=24pt (0.20 pts)")
                    total_score += 0.20
                    title_found = True
                elif is_centered and has_bold:
                    print("PARTIAL: Component 1 — Title found centered+bold but size not >=24pt (0.10 pts)")
                    total_score += 0.10
                    title_found = True
                elif 'certificate of achievement' in p.text.lower():
                    print("PARTIAL: Component 1 — Title text found but formatting incomplete (0.05 pts)")
                    total_score += 0.05
                    title_found = True
                break
        if not title_found:
            print("FAIL: Component 1 — 'CERTIFICATE OF ACHIEVEMENT' title not found")
    except Exception as e:
        print("ERROR: Component 1 — {}".format(e))

    # Component 2: "This is to certify that" + blank name line in ~20pt (0.15 points)
    try:
        certify_found = False
        name_line_found = False
        for i, p in enumerate(doc.paragraphs):
            if 'this is to certify that' in p.text.lower():
                certify_found = True
            # Check for a name placeholder line (underscores or blank) with ~20pt size
            if certify_found and not name_line_found:
                for r in p.runs:
                    if r.font.size and r.font.size.pt >= 18:
                        name_line_found = True
                        break
                # Also check next paragraphs for the 20pt line
                if not name_line_found and i + 1 < len(doc.paragraphs):
                    next_p = doc.paragraphs[i + 1]
                    for r in next_p.runs:
                        if r.font.size and r.font.size.pt >= 18:
                            name_line_found = True
                            break
        if certify_found and name_line_found:
            print("PASS: Component 2 — 'This is to certify that' + name line >=18pt (0.15 pts)")
            total_score += 0.15
        elif certify_found:
            print("PARTIAL: Component 2 — 'This is to certify that' found but no 20pt name line (0.07 pts)")
            total_score += 0.07
        else:
            print("FAIL: Component 2 — 'This is to certify that' not found")
    except Exception as e:
        print("ERROR: Component 2 — {}".format(e))

    # Component 3: "has successfully completed" + program name (0.15 points)
    try:
        completed_found = any('has successfully completed' in t for t in all_text_lower)
        program_found = any('advanced project management professional' in t for t in all_text_lower)
        if completed_found and program_found:
            print("PASS: Component 3 — 'has successfully completed' + program name found (0.15 pts)")
            total_score += 0.15
        elif completed_found or program_found:
            print("PARTIAL: Component 3 — only one of completion text or program name found (0.07 pts)")
            total_score += 0.07
        else:
            print("FAIL: Component 3 — neither completion text nor program name found")
    except Exception as e:
        print("ERROR: Component 3 — {}".format(e))

    # Component 4: Date of completion and certificate number fields (0.15 points)
    try:
        date_found = any('date' in t and 'completion' in t for t in all_text_lower)
        cert_num_found = any('certificate' in t and ('no' in t or 'number' in t or '#' in t) for t in all_text_lower)
        if date_found and cert_num_found:
            print("PASS: Component 4 — Date of completion + certificate number fields found (0.15 pts)")
            total_score += 0.15
        elif date_found or cert_num_found:
            print("PARTIAL: Component 4 — only one of date/cert number found (0.07 pts)")
            total_score += 0.07
        else:
            print("FAIL: Component 4 — date and certificate number fields not found")
    except Exception as e:
        print("ERROR: Component 4 — {}".format(e))

    # Component 5: Signature table with Program Director and Dean (0.15 points)
    try:
        sig_table_found = False
        director_found = False
        dean_found = False
        for table in doc.tables:
            table_text = ''
            for row in table.rows:
                for cell in row.cells:
                    cell_lower = cell.text.lower().strip()
                    table_text += cell_lower + ' '
                    if 'program director' in cell_lower or 'director' in cell_lower:
                        director_found = True
                    if 'dean' in cell_lower:
                        dean_found = True
            if director_found or dean_found:
                sig_table_found = True

        # Also check if signature lines are in paragraphs (not table)
        if not sig_table_found:
            for t in all_text_lower:
                if 'program director' in t or 'director' in t:
                    director_found = True
                if 'dean' in t:
                    dean_found = True

        if director_found and dean_found:
            print("PASS: Component 5 — Signature lines for Program Director and Dean found (0.15 pts)")
            total_score += 0.15
        elif director_found or dean_found:
            print("PARTIAL: Component 5 — only one signature role found (0.07 pts)")
            total_score += 0.07
        else:
            print("FAIL: Component 5 — signature lines not found")
    except Exception as e:
        print("ERROR: Component 5 — {}".format(e))

    # Component 6: Institution name "National Institute of Management" (0.10 points)
    try:
        inst_found = any('national institute of management' in t for t in all_text_lower)
        if inst_found:
            print("PASS: Component 6 — 'National Institute of Management' found (0.10 pts)")
            total_score += 0.10
        else:
            print("FAIL: Component 6 — 'National Institute of Management' not found")
    except Exception as e:
        print("ERROR: Component 6 — {}".format(e))

    # Component 7: Page border present (0.10 points)
    try:
        section = doc.sections[0]
        sectPr = section._sectPr
        pgBorders = sectPr.findall(qn('w:pgBorders'))
        if pgBorders and len(pgBorders) > 0:
            # Verify at least top and bottom borders exist
            border_elem = pgBorders[0]
            sides = [child.tag.split('}')[-1] for child in border_elem]
            has_top = 'top' in sides
            has_bottom = 'bottom' in sides
            if has_top and has_bottom:
                print("PASS: Component 7 — Page border present with top+bottom (0.10 pts)")
                total_score += 0.10
            else:
                print("PARTIAL: Component 7 — pgBorders element exists but incomplete (0.05 pts)")
                total_score += 0.05
        else:
            print("FAIL: Component 7 — No page border found")
    except Exception as e:
        print("ERROR: Component 7 — {}".format(e))

    final_score = round(min(total_score, 1.0), 2)
    print("")
    print("Score: {}/1.0".format(final_score))
    print("REWARD: {}".format(final_score))
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = '{}/{}.docx'.format(WORKDIR, TASK_ID)
if not os.path.exists(file_path):
    print("File not found: {}".format(file_path))
    print("REWARD: 0.0")
else:
    verify_task(file_path)
