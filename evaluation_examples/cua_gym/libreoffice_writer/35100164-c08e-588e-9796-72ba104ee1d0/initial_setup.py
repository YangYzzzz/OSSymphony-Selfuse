"""
Initial Setup: Training manual with plain bold text (no heading styles applied)
Task ID: osworld_writer_heading_styles_004
Domain: libreoffice_writer

Creates a training manual document where the title, chapter titles, and sub-topics
are all formatted as plain bold text of varying sizes — no heading styles are applied.
The agent must apply heading hierarchy (H1, H2, H3) and update Heading 2 to 14pt.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_heading_styles_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Document structure:
    # Title (1): "Employee Onboarding Training Manual" — plain bold, 20pt
    # Chapter 1: "Chapter 1: Company Overview" — plain bold, 16pt
    #   Sub-topic 1.1: "Company History and Mission" — plain bold, 12pt
    #   Sub-topic 1.2: "Organizational Structure" — plain bold, 12pt
    # Chapter 2: "Chapter 2: Workplace Policies" — plain bold, 16pt
    #   Sub-topic 2.1: "Code of Conduct" — plain bold, 12pt
    #   Sub-topic 2.2: "Attendance and Leave Policy" — plain bold, 12pt
    # Chapter 3: "Chapter 3: Benefits and Compensation" — plain bold, 16pt
    #   Sub-topic 3.1: "Health and Wellness Benefits" — plain bold, 12pt
    #   Sub-topic 3.2: "Salary Review Process" — plain bold, 12pt
    # Chapter 4: "Chapter 4: Professional Development" — plain bold, 16pt
    #   Sub-topic 4.1: "Training and Certification Programs" — plain bold, 12pt
    #   Sub-topic 4.2: "Performance Evaluation Framework" — plain bold, 12pt
    # All are Normal style (no heading styles), just manually bolded/sized

    # --- Title ---
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Employee Onboarding Training Manual")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        "Welcome to the organization. This manual provides essential information "
        "to help new employees understand company culture, policies, and expectations. "
        "Please read through each section carefully and consult your manager with any questions."
    )

    # --- Chapter 1 ---
    ch1_para = doc.add_paragraph()
    ch1_run = ch1_para.add_run("Chapter 1: Company Overview")
    ch1_run.bold = True
    ch1_run.font.size = Pt(16)
    ch1_para.paragraph_format.space_before = Pt(12)
    ch1_para.paragraph_format.space_after = Pt(6)

    # Sub-topic 1.1
    s11_para = doc.add_paragraph()
    s11_run = s11_para.add_run("Company History and Mission")
    s11_run.bold = True
    s11_run.font.size = Pt(12)
    s11_para.paragraph_format.space_before = Pt(6)

    doc.add_paragraph(
        "Founded in 1998, the company has grown from a small startup to a global leader "
        "in enterprise software solutions. Our mission is to empower organizations through "
        "innovative technology and exceptional customer service. Over the past two decades, "
        "we have expanded operations to 35 countries and serve more than 10,000 clients worldwide."
    )

    # Sub-topic 1.2
    s12_para = doc.add_paragraph()
    s12_run = s12_para.add_run("Organizational Structure")
    s12_run.bold = True
    s12_run.font.size = Pt(12)
    s12_para.paragraph_format.space_before = Pt(6)

    doc.add_paragraph(
        "The company is organized into five core divisions: Engineering, Sales, Marketing, "
        "Customer Success, and Operations. Each division is led by a Vice President who reports "
        "to the Chief Executive Officer. Regional offices in North America, Europe, and Asia-Pacific "
        "each have their own General Manager responsible for local business performance."
    )

    # --- Chapter 2 ---
    ch2_para = doc.add_paragraph()
    ch2_run = ch2_para.add_run("Chapter 2: Workplace Policies")
    ch2_run.bold = True
    ch2_run.font.size = Pt(16)
    ch2_para.paragraph_format.space_before = Pt(12)
    ch2_para.paragraph_format.space_after = Pt(6)

    # Sub-topic 2.1
    s21_para = doc.add_paragraph()
    s21_run = s21_para.add_run("Code of Conduct")
    s21_run.bold = True
    s21_run.font.size = Pt(12)
    s21_para.paragraph_format.space_before = Pt(6)

    doc.add_paragraph(
        "All employees are expected to maintain the highest standards of professional conduct. "
        "This includes treating colleagues, clients, and partners with respect and dignity, "
        "adhering to confidentiality agreements, avoiding conflicts of interest, and reporting "
        "any suspected violations of company policy to the Human Resources department promptly."
    )

    # Sub-topic 2.2
    s22_para = doc.add_paragraph()
    s22_run = s22_para.add_run("Attendance and Leave Policy")
    s22_run.bold = True
    s22_run.font.size = Pt(12)
    s22_para.paragraph_format.space_before = Pt(6)

    doc.add_paragraph(
        "Standard working hours are 9:00 AM to 6:00 PM, Monday through Friday. Employees receive "
        "15 days of annual leave, 10 sick days, and 3 personal days per calendar year. Leave requests "
        "must be submitted at least two weeks in advance through the HR portal. Unpaid leave may be "
        "granted at the discretion of the department head for extended absences."
    )

    # --- Chapter 3 ---
    ch3_para = doc.add_paragraph()
    ch3_run = ch3_para.add_run("Chapter 3: Benefits and Compensation")
    ch3_run.bold = True
    ch3_run.font.size = Pt(16)
    ch3_para.paragraph_format.space_before = Pt(12)
    ch3_para.paragraph_format.space_after = Pt(6)

    # Sub-topic 3.1
    s31_para = doc.add_paragraph()
    s31_run = s31_para.add_run("Health and Wellness Benefits")
    s31_run.bold = True
    s31_run.font.size = Pt(12)
    s31_para.paragraph_format.space_before = Pt(6)

    doc.add_paragraph(
        "The company offers a comprehensive health benefits package including medical, dental, and vision "
        "coverage for employees and their immediate family members. Additionally, employees have access to "
        "an Employee Assistance Program (EAP) providing confidential counseling services, a $500 annual "
        "wellness stipend for gym memberships or fitness equipment, and on-site mental health resources."
    )

    # Sub-topic 3.2
    s32_para = doc.add_paragraph()
    s32_run = s32_para.add_run("Salary Review Process")
    s32_run.bold = True
    s32_run.font.size = Pt(12)
    s32_para.paragraph_format.space_before = Pt(6)

    doc.add_paragraph(
        "Salary reviews are conducted annually in December, with adjustments taking effect January 1st. "
        "The review process considers individual performance ratings, market benchmarking data, and "
        "departmental budget allocations. Employees rated 'Exceeds Expectations' typically receive "
        "merit increases of 4-8%, while 'Meets Expectations' ratings correspond to increases of 2-4%."
    )

    # --- Chapter 4 ---
    ch4_para = doc.add_paragraph()
    ch4_run = ch4_para.add_run("Chapter 4: Professional Development")
    ch4_run.bold = True
    ch4_run.font.size = Pt(16)
    ch4_para.paragraph_format.space_before = Pt(12)
    ch4_para.paragraph_format.space_after = Pt(6)

    # Sub-topic 4.1
    s41_para = doc.add_paragraph()
    s41_run = s41_para.add_run("Training and Certification Programs")
    s41_run.bold = True
    s41_run.font.size = Pt(12)
    s41_para.paragraph_format.space_before = Pt(6)

    doc.add_paragraph(
        "The company invests in employee growth through a robust learning and development program. "
        "New hires participate in a structured 90-day onboarding curriculum, and all employees have "
        "access to over 2,000 online courses via the internal learning management system. An annual "
        "professional development budget of $1,500 per employee covers external conferences, workshops, "
        "and industry certification exam fees."
    )

    # Sub-topic 4.2
    s42_para = doc.add_paragraph()
    s42_run = s42_para.add_run("Performance Evaluation Framework")
    s42_run.bold = True
    s42_run.font.size = Pt(12)
    s42_para.paragraph_format.space_before = Pt(6)

    doc.add_paragraph(
        "Performance evaluations are conducted twice annually: a mid-year check-in in June and a "
        "comprehensive year-end review in December. Evaluations assess achievement of SMART goals, "
        "competency demonstration across five core areas, and 360-degree feedback from managers, "
        "peers, and direct reports. Results are used to inform promotion decisions, development plans, "
        "and compensation adjustments for the following year."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
