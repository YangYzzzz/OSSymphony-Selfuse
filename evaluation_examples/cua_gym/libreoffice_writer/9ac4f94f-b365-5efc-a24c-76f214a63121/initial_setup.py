"""
Initial Setup: Writer document with inline image (no text wrap)
Task ID: writer_fs_026
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_026'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
IMG_PATH = f'{WORKDIR}/sample_image.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_sample_image():
    """Create a simple colored image to insert into the document."""
    from PIL import Image, ImageDraw, ImageFont
    img = Image.new('RGB', (400, 300), color=(70, 130, 180))
    draw = ImageDraw.Draw(img)
    # Draw some shapes to make it look like a chart/figure
    draw.rectangle([30, 40, 370, 260], outline='white', width=2)
    draw.rectangle([50, 180, 100, 260], fill=(46, 204, 113))
    draw.rectangle([120, 140, 170, 260], fill=(52, 152, 219))
    draw.rectangle([190, 100, 240, 260], fill=(155, 89, 182))
    draw.rectangle([260, 160, 310, 260], fill=(231, 76, 60))
    draw.rectangle([330, 120, 370, 260], fill=(241, 196, 15))
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 18)
    except Exception:
        font = ImageFont.load_default()
    draw.text((100, 10), "Q3 Sales Report", fill='white', font=font)
    img.save(IMG_PATH)
    print(f'Sample image created: {IMG_PATH}')


def create_initial():
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_heading('Quarterly Business Performance Report', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by the Analytics Division — September 2025')
    run.font.size = Pt(11)
    run.italic = True

    doc.add_paragraph('')  # spacer

    # Intro paragraph
    p1 = doc.add_paragraph()
    run = p1.add_run(
        'The third quarter of fiscal year 2025 has demonstrated remarkable growth across '
        'all key performance indicators. Revenue increased by 18.3% compared to Q2, driven '
        'primarily by strong demand in the enterprise software segment and expanded partnerships '
        'in the Asia-Pacific region. Customer acquisition costs decreased by 7.2%, reflecting '
        'improved marketing efficiency and organic referral growth.'
    )
    run.font.size = Pt(11)

    # Second paragraph before image
    p2 = doc.add_paragraph()
    run = p2.add_run(
        'The following chart illustrates the comparative sales performance across our five '
        'primary product lines during Q3. Each bar represents total revenue in millions of '
        'dollars, with the enterprise solutions segment continuing to lead overall growth.'
    )
    run.font.size = Pt(11)

    # Insert image inline (default "None" text wrap)
    doc.add_picture(IMG_PATH, width=Inches(4.0))

    # More body text after image
    p3 = doc.add_paragraph()
    run = p3.add_run(
        'As evidenced by the data above, the Enterprise Solutions division generated $4.8M in '
        'revenue, representing a 22% year-over-year increase. The Cloud Services team also '
        'posted impressive numbers at $3.9M, benefiting from the migration of legacy clients '
        'to our new SaaS platform. Meanwhile, the Professional Services and Hardware divisions '
        'maintained steady performance, contributing $2.1M and $1.7M respectively.'
    )
    run.font.size = Pt(11)

    p4 = doc.add_paragraph()
    run = p4.add_run(
        'Looking ahead to Q4, the leadership team has identified three strategic priorities: '
        '(1) accelerating the rollout of Version 5.0 of our flagship product, (2) expanding '
        'the partner ecosystem in Europe and Latin America, and (3) investing in AI-driven '
        'analytics capabilities to enhance customer retention. The projected revenue target '
        'for Q4 is $14.2M, representing a 12% increase over Q3 performance.'
    )
    run.font.size = Pt(11)

    p5 = doc.add_paragraph()
    run = p5.add_run(
        'The finance team has also noted a significant improvement in operating margins, '
        'which rose from 23.1% in Q2 to 26.8% in Q3. This improvement is attributed to '
        'reduced infrastructure costs following the data center consolidation completed in '
        'July, as well as headcount optimization in the support division. Net income for '
        'the quarter reached $3.2M, exceeding analyst expectations by approximately 8%.'
    )
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_sample_image()
create_initial()
