"""
Initial Setup: Programming tutorial document with code snippets in default style
Task ID: writer_rd_027
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_027'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font for body text
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Title ---
    heading = doc.add_heading('Introduction to Python Data Processing', level=1)

    # --- Intro paragraph ---
    p = doc.add_paragraph(
        'Python has become the go-to language for data analysis and processing tasks. '
        'Its rich ecosystem of libraries like pandas, NumPy, and matplotlib make it possible '
        'to handle complex data workflows with relatively simple code. In this tutorial, we will '
        'walk through three common data processing patterns that every Python developer should know.'
    )

    # --- Section 1 ---
    doc.add_heading('1. Reading and Filtering CSV Data', level=2)

    p = doc.add_paragraph(
        'The first step in most data pipelines is loading raw data from files. The pandas library '
        'provides a straightforward API for reading CSV files into DataFrames. Once loaded, you can '
        'filter rows based on conditions, select specific columns, and transform values. Here is a '
        'basic example that reads a sales report and filters for high-value transactions:'
    )

    # Code snippet 1 - in default paragraph style (no special formatting)
    code1 = doc.add_paragraph(
        'import pandas as pd\n'
        '\n'
        'df = pd.read_csv("sales_report_2025.csv")\n'
        'high_value = df[df["amount"] > 10000]\n'
        'high_value = high_value.sort_values("date", ascending=False)\n'
        'print(f"Found {len(high_value)} high-value transactions")\n'
        'high_value.to_csv("filtered_sales.csv", index=False)'
    )

    p = doc.add_paragraph(
        'This approach works well for moderate-sized files. For datasets larger than available memory, '
        'consider using the chunksize parameter or switching to Dask for out-of-core processing.'
    )

    # --- Section 2 ---
    doc.add_heading('2. Aggregating Data with GroupBy', level=2)

    p = doc.add_paragraph(
        'Grouping and aggregating data is essential for generating summary statistics. The groupby '
        'method in pandas splits the DataFrame into groups based on one or more columns, applies an '
        'aggregation function, and combines the results. This pattern is similar to SQL GROUP BY '
        'clauses. The following example computes quarterly revenue by region:'
    )

    # Code snippet 2 - in default paragraph style
    code2 = doc.add_paragraph(
        'quarterly = df.groupby(["region", "quarter"]).agg(\n'
        '    total_revenue=("revenue", "sum"),\n'
        '    avg_order_size=("order_amount", "mean"),\n'
        '    transaction_count=("order_id", "count")\n'
        ').reset_index()\n'
        '\n'
        'quarterly["growth_rate"] = quarterly.groupby("region")["total_revenue"].pct_change()\n'
        'print(quarterly.head(10))'
    )

    p = doc.add_paragraph(
        'Notice how we chain the pct_change method to compute growth rates within each region. '
        'This is a powerful technique for time-series analysis that avoids explicit looping.'
    )

    # --- Section 3 ---
    doc.add_heading('3. Merging Multiple Data Sources', level=2)

    p = doc.add_paragraph(
        'Real-world analysis often requires combining data from multiple sources. The merge function '
        'in pandas performs SQL-style joins between DataFrames. You can specify the join type (inner, '
        'left, right, outer) and the columns to join on. Below is an example that combines customer '
        'profiles with their purchase history and product details:'
    )

    # Code snippet 3 - in default paragraph style
    code3 = doc.add_paragraph(
        'customers = pd.read_csv("customers.csv")\n'
        'purchases = pd.read_csv("purchase_history.csv")\n'
        'products = pd.read_csv("product_catalog.csv")\n'
        '\n'
        'merged = purchases.merge(customers, on="customer_id", how="left")\n'
        'merged = merged.merge(products, on="product_id", how="left")\n'
        'merged["total_spent"] = merged["quantity"] * merged["unit_price"]\n'
        'summary = merged.groupby("customer_name")["total_spent"].sum().sort_values(ascending=False)\n'
        'print(summary.head(20))'
    )

    # --- Conclusion ---
    doc.add_heading('Conclusion', level=2)

    p = doc.add_paragraph(
        'These three patterns -- reading/filtering, aggregating, and merging -- form the foundation '
        'of most data processing workflows in Python. By combining them creatively, you can build '
        'sophisticated analysis pipelines that handle real business requirements. As your datasets grow, '
        'consider profiling your code with tools like line_profiler and optimizing bottlenecks with '
        'vectorized operations rather than row-by-row iteration.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
