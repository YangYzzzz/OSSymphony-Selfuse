"""
Reward Script: Create 'Code Block' paragraph style and apply to code snippets
Task ID: writer_rd_027
Domain: libreoffice_writer
Scoring:
  C1 (0.20) — 'Code Block' custom paragraph style exists
  C2 (0.20) — Style font is Courier New 10pt
  C3 (0.15) — Style paragraph shading fill is #F2F2F2
  C4 (0.15) — Style left/right indent ~0.5 cm
  C5 (0.10) — Style space before/after ~0.3 cm
  C6 (0.20) — All 3 code snippet paragraphs use 'Code Block' style
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_027'


def persist_app_state(domain: str):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------------------------------------------------------------
    # Component 1: 'Code Block' custom paragraph style exists (0.20)
    # ---------------------------------------------------------------
    code_block_style = None
    try:
        code_block_style = doc.styles['Code Block']
        if code_block_style is not None and not code_block_style.builtin:
            print(f"PASS: C1 — 'Code Block' custom paragraph style exists (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: C1 — 'Code Block' style is builtin or missing")
    except KeyError:
        print(f"FAIL: C1 — 'Code Block' style not found in document styles")
    except Exception as e:
        print(f"ERROR: C1 — {e}")

    if code_block_style is None:
        # No point checking style properties if style doesn't exist
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # ---------------------------------------------------------------
    # Component 2: Style font is Courier New 10pt (0.20)
    # ---------------------------------------------------------------
    try:
        font = code_block_style.font
        font_name = font.name
        font_size_pt = font.size.pt if font.size else None

        # Also check XML directly for robustness
        rPr = code_block_style.element.find(qn('w:rPr'))
        xml_font_name = None
        xml_font_size = None
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                xml_font_name = rFonts.get(qn('w:ascii'))
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                # w:sz is in half-points
                xml_font_size = int(sz.get(qn('w:val'))) / 2.0

        actual_name = font_name or xml_font_name
        actual_size = font_size_pt or xml_font_size

        name_ok = actual_name is not None and 'courier' in actual_name.lower()
        size_ok = actual_size is not None and abs(actual_size - 10.0) < 0.5

        if name_ok and size_ok:
            print(f"PASS: C2 — Font is '{actual_name}' {actual_size}pt (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: C2 — Expected Courier New 10pt, found '{actual_name}' {actual_size}pt")
    except Exception as e:
        print(f"ERROR: C2 — {e}")

    # ---------------------------------------------------------------
    # Component 3: Style paragraph shading fill #F2F2F2 (0.15)
    # ---------------------------------------------------------------
    try:
        pPr = code_block_style.element.find(qn('w:pPr'))
        shd = pPr.find(qn('w:shd')) if pPr is not None else None
        fill_val = shd.get(qn('w:fill')).upper() if shd is not None else None

        if fill_val == 'F2F2F2':
            print(f"PASS: C3 — Paragraph shading fill is #{fill_val} (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: C3 — Expected shading fill #F2F2F2, found: {fill_val}")
    except Exception as e:
        print(f"ERROR: C3 — {e}")

    # ---------------------------------------------------------------
    # Component 4: Style left/right indent ~0.5 cm (0.15)
    # ---------------------------------------------------------------
    try:
        pf = code_block_style.paragraph_format
        left_cm = pf.left_indent / Cm(1) if pf.left_indent else 0
        right_cm = pf.right_indent / Cm(1) if pf.right_indent else 0

        # Tolerance: within 0.1 cm of 0.5
        left_ok = abs(left_cm - 0.5) < 0.1
        right_ok = abs(right_cm - 0.5) < 0.1

        if left_ok and right_ok:
            print(f"PASS: C4 — Left indent={left_cm:.3f}cm, Right indent={right_cm:.3f}cm (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: C4 — Expected ~0.5cm L/R indent, found L={left_cm:.3f}cm R={right_cm:.3f}cm")
    except Exception as e:
        print(f"ERROR: C4 — {e}")

    # ---------------------------------------------------------------
    # Component 5: Style space before/after ~0.3 cm (0.10)
    # ---------------------------------------------------------------
    try:
        pf = code_block_style.paragraph_format
        before_cm = pf.space_before / Cm(1) if pf.space_before else 0
        after_cm = pf.space_after / Cm(1) if pf.space_after else 0

        # Tolerance: within 0.1 cm of 0.3
        before_ok = abs(before_cm - 0.3) < 0.1
        after_ok = abs(after_cm - 0.3) < 0.1

        if before_ok and after_ok:
            print(f"PASS: C5 — Space before={before_cm:.3f}cm, after={after_cm:.3f}cm (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: C5 — Expected ~0.3cm spacing, found before={before_cm:.3f}cm after={after_cm:.3f}cm")
    except Exception as e:
        print(f"ERROR: C5 — {e}")

    # ---------------------------------------------------------------
    # Component 6: All 3 code snippet paragraphs use 'Code Block' style (0.20)
    # ---------------------------------------------------------------
    try:
        # Identify code snippet paragraphs by content patterns
        # They contain Python code (import, df., pd., etc.)
        code_indicators = ['import pandas', 'df.groupby', 'pd.read_csv', 'customers =', 'purchases =']
        code_para_indices = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if any(ind in text for ind in code_indicators):
                code_para_indices.append(i)

        if len(code_para_indices) < 3:
            # Fallback: check all paragraphs with 'Code Block' style
            code_block_paras = [i for i, p in enumerate(doc.paragraphs)
                                if p.style and p.style.name == 'Code Block']
            if len(code_block_paras) >= 3:
                print(f"PASS: C6 — {len(code_block_paras)} paragraphs use 'Code Block' style (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: C6 — Only {len(code_block_paras)} paragraphs use 'Code Block' style, expected >= 3")
        else:
            # Check that identified code paragraphs have 'Code Block' style
            applied_count = 0
            for idx in code_para_indices:
                para = doc.paragraphs[idx]
                style_name = para.style.name if para.style else 'None'
                if style_name == 'Code Block':
                    applied_count += 1
                else:
                    print(f"  INFO: P{idx} has style '{style_name}' instead of 'Code Block'")

            if applied_count >= 3:
                print(f"PASS: C6 — {applied_count}/{len(code_para_indices)} code paragraphs use 'Code Block' (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: C6 — Only {applied_count}/{len(code_para_indices)} code paragraphs use 'Code Block'")
    except Exception as e:
        print(f"ERROR: C6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
