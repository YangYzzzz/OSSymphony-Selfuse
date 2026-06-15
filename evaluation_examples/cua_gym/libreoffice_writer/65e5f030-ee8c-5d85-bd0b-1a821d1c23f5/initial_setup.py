"""
Initial Setup: Client intake form with labels but no form controls
Task ID: writer_biz_064
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_064'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    heading = doc.add_heading('New Client Intake Form', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle / instructions
    intro = doc.add_paragraph()
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = intro.add_run('Please complete all fields below for new client registration. '
                        'All information provided will be kept confidential in accordance '
                        'with our data privacy policy.')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    doc.add_paragraph()  # spacer

    # --- Section: Company Information ---
    sec1 = doc.add_heading('Company Information', level=2)

    # Company Name label
    p1 = doc.add_paragraph()
    r1 = p1.add_run('Company Name:')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.name = 'Calibri'
    p1.add_run('  ')  # space for future form field

    doc.add_paragraph()  # spacer

    # Industry label
    p2 = doc.add_paragraph()
    r2 = p2.add_run('Industry:')
    r2.bold = True
    r2.font.size = Pt(12)
    r2.font.name = 'Calibri'
    p2.add_run('  ')

    doc.add_paragraph()  # spacer

    # --- Section: Contact Details ---
    sec2 = doc.add_heading('Contact Details', level=2)

    # Contact Person label
    p3 = doc.add_paragraph()
    r3 = p3.add_run('Contact Person:')
    r3.bold = True
    r3.font.size = Pt(12)
    r3.font.name = 'Calibri'
    p3.add_run('  ')

    doc.add_paragraph()  # spacer

    # Phone label
    p4 = doc.add_paragraph()
    r4 = p4.add_run('Phone:')
    r4.bold = True
    r4.font.size = Pt(12)
    r4.font.name = 'Calibri'
    p4.add_run('  ')

    doc.add_paragraph()  # spacer

    # Email label
    p5 = doc.add_paragraph()
    r5 = p5.add_run('Email:')
    r5.bold = True
    r5.font.size = Pt(12)
    r5.font.name = 'Calibri'
    p5.add_run('  ')

    doc.add_paragraph()  # spacer

    # --- Section: Agreement ---
    sec3 = doc.add_heading('Terms and Conditions', level=2)

    agree_text = doc.add_paragraph()
    r_agree = agree_text.add_run(
        'By checking the box below, you confirm that all information provided '
        'is accurate and that you have read and agree to our Terms and Conditions, '
        'including our data processing agreement and service level commitments.'
    )
    r_agree.font.size = Pt(11)
    r_agree.font.name = 'Calibri'

    doc.add_paragraph()  # spacer

    # Accepts Terms label
    p6 = doc.add_paragraph()
    r6 = p6.add_run('Accepts Terms and Conditions:')
    r6.bold = True
    r6.font.size = Pt(12)
    r6.font.name = 'Calibri'
    p6.add_run('  ')

    doc.add_paragraph()  # spacer

    # --- Footer note ---
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r_footer = footer_para.add_run('Meridian Consulting Group — Client Services Department')
    r_footer.font.size = Pt(9)
    r_footer.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    r_footer.font.name = 'Calibri'
    r_footer.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
