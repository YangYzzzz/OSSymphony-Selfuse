"""
Reward Script: Wedding program in LibreOffice Writer
Task ID: writer_wf_059
Domain: libreoffice_writer
Scoring:
  C1: Title 'Emma & Michael' centered, 22pt, bold (0.15)
  C2: Wedding date centered, 14pt, italic (0.10)
  C3: Ceremony Order section with 8 numbered items (0.20)
  C4: Wedding Party table with header + 8 data rows, 2 cols (0.20)
  C5: Readings section with 2 reading titles + readers (0.10)
  C6: Special Thanks paragraph present (0.05)
  C7: Reception Details section present (0.10)
  C8: Decorative page border present (0.10)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_059'


def verify_task(file_path):
    """Verify wedding program task completion with progressive scoring."""
    total_score = 0.0

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError as e:
        print(f"CRITICAL: Missing library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_paras = doc.paragraphs
    all_text = [p.text.strip() for p in all_paras]

    # Helper: find paragraph containing text (case-insensitive)
    def find_para_containing(keyword):
        for p in all_paras:
            if keyword.lower() in p.text.lower():
                return p
        return None

    # Component 1: Title 'Emma & Michael' centered, 22pt, bold (0.15 points)
    try:
        title_para = find_para_containing('emma')
        if title_para and 'michael' in title_para.text.lower():
            is_centered = (title_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
            has_bold = any(r.bold for r in title_para.runs if r.text.strip())
            has_22pt = any(
                r.font.size is not None and abs(r.font.size.pt - 22.0) < 1.0
                for r in title_para.runs if r.text.strip()
            )
            if is_centered and has_bold and has_22pt:
                print(f"PASS: Component 1 — Title 'Emma & Michael' centered, 22pt, bold (0.15 pts)")
                total_score += 0.15
            elif has_bold or has_22pt:
                partial = 0.05
                print(f"PARTIAL: Component 1 — Title found but missing some formatting (centered={is_centered}, bold={has_bold}, 22pt={has_22pt}) ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 1 — Title found but no formatting (centered={is_centered}, bold={has_bold}, 22pt={has_22pt})")
        else:
            print("FAIL: Component 1 — Title containing 'Emma' and 'Michael' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Wedding date centered, 14pt, italic (0.10 points)
    try:
        date_found = False
        for p in all_paras:
            text = p.text.strip()
            # Look for a date-like paragraph (contains year or month or 'june', 'saturday', etc.)
            if re.search(r'(20\d{2}|january|february|march|april|may|june|july|august|september|october|november|december|saturday|sunday|monday|tuesday|wednesday|thursday|friday)', text, re.IGNORECASE):
                if 'emma' not in text.lower():  # Skip the title
                    is_centered = (p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
                    has_italic = any(r.italic for r in p.runs if r.text.strip())
                    has_14pt = any(
                        r.font.size is not None and abs(r.font.size.pt - 14.0) < 1.0
                        for r in p.runs if r.text.strip()
                    )
                    if is_centered and has_italic and has_14pt:
                        print(f"PASS: Component 2 — Date '{text}' centered, 14pt, italic (0.10 pts)")
                        total_score += 0.10
                        date_found = True
                        break
                    elif has_italic or has_14pt:
                        partial = 0.03
                        print(f"PARTIAL: Component 2 — Date found but missing formatting (centered={is_centered}, italic={has_italic}, 14pt={has_14pt}) ({partial} pts)")
                        total_score += partial
                        date_found = True
                        break
        if not date_found:
            print("FAIL: Component 2 — Wedding date paragraph not found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Ceremony Order section with 8 numbered items (0.20 points)
    try:
        ceremony_heading = find_para_containing('ceremony')
        if ceremony_heading:
            # Count numbered items after the ceremony heading
            ceremony_idx = all_paras.index(ceremony_heading)
            numbered_items = []
            for p in all_paras[ceremony_idx + 1:]:
                text = p.text.strip()
                if not text:
                    continue
                # Stop at next section heading or decorative separator
                if any(kw in text.lower() for kw in ['wedding party', 'readings', 'special thanks', 'reception']):
                    break
                # Check for numbered pattern: "1.", "2.", etc.
                if re.match(r'^\d+[\.\)]\s', text):
                    numbered_items.append(text)
                # Also check list-style paragraphs
                elif p.style.name.startswith('List'):
                    numbered_items.append(text)

            num_items = len(numbered_items)
            if num_items >= 8:
                print(f"PASS: Component 3 — Ceremony Order with {num_items} numbered items (0.20 pts)")
                total_score += 0.20
            elif num_items >= 4:
                partial = 0.10
                print(f"PARTIAL: Component 3 — Ceremony Order with {num_items}/8 items ({partial} pts)")
                total_score += partial
            elif num_items > 0:
                partial = 0.05
                print(f"PARTIAL: Component 3 — Ceremony Order with {num_items}/8 items ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — Ceremony Order heading found but no numbered items")
        else:
            print("FAIL: Component 3 — 'Ceremony Order' section not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Wedding Party table with header + 8 data rows, 2 cols (0.20 points)
    try:
        if len(doc.tables) > 0:
            # Find the table that looks like a wedding party table (has Role/Name headers or similar)
            best_table = None
            for t in doc.tables:
                header_texts = [c.text.strip().lower() for c in t.rows[0].cells]
                if 'role' in header_texts or 'name' in header_texts:
                    best_table = t
                    break
            if best_table is None:
                best_table = doc.tables[0]  # fallback to first table

            num_rows = len(best_table.rows)
            num_cols = len(best_table.columns)
            has_header = any('role' in c.text.strip().lower() or 'name' in c.text.strip().lower()
                           for c in best_table.rows[0].cells)

            # We expect header + 8 data rows = 9 rows total, 2 columns
            if num_cols >= 2 and num_rows >= 9 and has_header:
                print(f"PASS: Component 4 — Wedding Party table: {num_rows} rows x {num_cols} cols with header (0.20 pts)")
                total_score += 0.20
            elif num_cols >= 2 and num_rows >= 5:
                partial = 0.10
                print(f"PARTIAL: Component 4 — Table found but only {num_rows} rows (expected 9) ({partial} pts)")
                total_score += partial
            elif num_cols >= 2:
                partial = 0.05
                print(f"PARTIAL: Component 4 — Table found with {num_rows} rows, {num_cols} cols ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — Table found but only {num_cols} columns (expected 2+)")
        else:
            print("FAIL: Component 4 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Readings section with 2 reading titles and readers (0.10 points)
    try:
        readings_heading = find_para_containing('readings')
        # Exclude paragraphs that are about 'scripture reading' in ceremony order
        readings_section_found = False
        for p in all_paras:
            text = p.text.strip().lower()
            if text == 'readings' or (text.startswith('readings') and 'scripture' not in text and len(text) < 20):
                readings_heading = p
                readings_section_found = True
                break

        if readings_section_found and readings_heading:
            readings_idx = all_paras.index(readings_heading)
            readings_content = []
            for p in all_paras[readings_idx + 1:]:
                text = p.text.strip()
                if not text or text == '— ✦ —':
                    if readings_content:
                        break
                    continue
                if any(kw in text.lower() for kw in ['special thanks', 'reception', 'wedding party']):
                    break
                readings_content.append(text)

            # Check for at least 2 reading titles and 2 reader names
            reading_titles = [t for t in readings_content if 'read by' not in t.lower()]
            reader_entries = [t for t in readings_content if 'read by' in t.lower()]

            if len(reading_titles) >= 2 and len(reader_entries) >= 2:
                print(f"PASS: Component 5 — Readings: {len(reading_titles)} titles, {len(reader_entries)} readers (0.10 pts)")
                total_score += 0.10
            elif len(reading_titles) >= 1 or len(reader_entries) >= 1:
                partial = 0.05
                print(f"PARTIAL: Component 5 — Readings incomplete: {len(reading_titles)} titles, {len(reader_entries)} readers ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 — Readings section found but no titles/readers: {readings_content}")
        else:
            print("FAIL: Component 5 — 'Readings' section heading not found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Special Thanks paragraph present (0.05 points)
    try:
        thanks_heading = find_para_containing('special thanks')
        if thanks_heading:
            # Look for content after the heading
            thanks_idx = all_paras.index(thanks_heading)
            has_content = False
            for p in all_paras[thanks_idx + 1:]:
                text = p.text.strip()
                if not text or text == '— ✦ —':
                    if has_content:
                        break
                    continue
                if any(kw in text.lower() for kw in ['reception', 'readings', 'ceremony']):
                    break
                if len(text) > 20:  # Substantive paragraph
                    has_content = True

            if has_content:
                print(f"PASS: Component 6 — Special Thanks section with content (0.05 pts)")
                total_score += 0.05
            else:
                # Still give credit for the heading
                partial = 0.02
                print(f"PARTIAL: Component 6 — Special Thanks heading found but no content ({partial} pts)")
                total_score += partial
        else:
            print("FAIL: Component 6 — 'Special Thanks' section not found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Reception Details section present (0.10 points)
    try:
        reception_heading = find_para_containing('reception')
        if reception_heading:
            reception_idx = all_paras.index(reception_heading)
            reception_content = []
            for p in all_paras[reception_idx + 1:]:
                text = p.text.strip()
                if not text:
                    continue
                if text == '— ✦ —':
                    continue
                reception_content.append(text)

            if len(reception_content) >= 2:
                print(f"PASS: Component 7 — Reception Details with {len(reception_content)} lines of info (0.10 pts)")
                total_score += 0.10
            elif len(reception_content) >= 1:
                partial = 0.05
                print(f"PARTIAL: Component 7 — Reception Details with only {len(reception_content)} line ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 7 — Reception heading found but no details")
        else:
            print("FAIL: Component 7 — 'Reception Details' section not found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # Component 8: Decorative page border present (0.10 points)
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        has_border = False
        for section in doc.sections:
            sectPr = section._sectPr
            borders = sectPr.findall('.//w:pgBorders', ns)
            if borders:
                # Verify it has at least top/bottom/left/right borders
                for pgBorders in borders:
                    sides = ['top', 'bottom', 'left', 'right']
                    found_sides = []
                    for side in sides:
                        border_el = pgBorders.find(f'w:{side}', ns)
                        if border_el is not None:
                            val = border_el.get(f'{{{ns["w"]}}}val', '')
                            if val and val != 'none':
                                found_sides.append(side)
                    if len(found_sides) == 4:
                        has_border = True
                        print(f"PASS: Component 8 — Page border on all 4 sides (0.10 pts)")
                        total_score += 0.10
                        break
                    elif len(found_sides) > 0:
                        partial = 0.05
                        print(f"PARTIAL: Component 8 — Page border on {len(found_sides)}/4 sides ({partial} pts)")
                        total_score += partial
                        has_border = True
                        break
            if has_border:
                break
        if not has_border:
            print("FAIL: Component 8 — No page border found in document")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
