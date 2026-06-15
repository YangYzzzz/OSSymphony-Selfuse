"""
Initial Setup: Create Annual Report with TOC using dot leaders
Task ID: writer_mt_063
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_063'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc_field(doc):
    """Add a TOC field that LibreOffice can update."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar_separate)

    run4 = paragraph.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run4._element.append(fldChar_end)

    return paragraph


def add_toc_entry(doc, text, page_num, level, tab_position):
    """Add a static TOC entry with dot leader."""
    # Use the appropriate TOC style
    style_name = f'TOC {level}'

    para = doc.add_paragraph(style=style_name)

    # Add the entry text
    run_text = para.add_run(text)

    # Add tab character
    run_tab = para.add_run('\t')

    # Add page number
    run_page = para.add_run(str(page_num))

    # Set tab stop with dot leader at right margin
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    return para


def ensure_toc_styles(doc, tab_position):
    """Ensure TOC 1, TOC 2, TOC 3 styles exist with dot leaders."""
    styles = doc.styles
    for level in range(1, 4):
        style_name = f'TOC {level}'
        try:
            style = styles[style_name]
        except KeyError:
            # Create TOC style if it doesn't exist
            style = styles.add_style(style_name, 1)  # WD_STYLE_TYPE.PARAGRAPH = 1

        # Set indentation based on level
        pf = style.paragraph_format
        pf.left_indent = Inches(0.2 * (level - 1))
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)

        # Set tab stop with dot leader
        pf.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        # Set font
        style.font.size = Pt(12 - (level - 1))
        style.font.name = 'Liberation Serif'


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    tab_position = Inches(6.5)  # Right-aligned tab near right margin

    # Ensure TOC styles exist with dot leaders
    ensure_toc_styles(doc, tab_position)

    # === Title Page ===
    title_para = doc.add_heading('Meridian Technologies Annual Report 2025', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Fiscal Year Ending December 31, 2025')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_page_break()

    # === Table of Contents ===
    toc_heading = doc.add_heading('Table of Contents', level=1)

    # Add static TOC entries with dot leaders (3 levels)
    # Level 1 entries
    add_toc_entry(doc, 'Executive Summary', 3, 1, tab_position)
    add_toc_entry(doc, 'Financial Highlights', 5, 1, tab_position)
    add_toc_entry(doc, 'Operations Review', 9, 1, tab_position)
    add_toc_entry(doc, 'Strategic Initiatives', 14, 1, tab_position)
    add_toc_entry(doc, 'Corporate Governance', 18, 1, tab_position)

    # Level 2 entries (indented)
    add_toc_entry(doc, 'Revenue Analysis', 5, 2, tab_position)
    add_toc_entry(doc, 'Cost Management', 6, 2, tab_position)
    add_toc_entry(doc, 'Balance Sheet Overview', 7, 2, tab_position)
    add_toc_entry(doc, 'Product Development', 9, 2, tab_position)
    add_toc_entry(doc, 'Market Expansion', 11, 2, tab_position)
    add_toc_entry(doc, 'Digital Transformation', 14, 2, tab_position)
    add_toc_entry(doc, 'Sustainability Goals', 16, 2, tab_position)

    # Level 3 entries (more indented)
    add_toc_entry(doc, 'Regional Revenue Breakdown', 5, 3, tab_position)
    add_toc_entry(doc, 'Year-over-Year Comparison', 6, 3, tab_position)
    add_toc_entry(doc, 'R&D Pipeline Status', 10, 3, tab_position)
    add_toc_entry(doc, 'Asia-Pacific Expansion', 12, 3, tab_position)
    add_toc_entry(doc, 'Cloud Infrastructure Roadmap', 15, 3, tab_position)

    doc.add_page_break()

    # === Executive Summary ===
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'Meridian Technologies delivered another year of strong performance in fiscal year 2025, '
        'achieving record revenue of $4.82 billion, representing a 14.3% increase over the previous '
        'fiscal year. Our strategic investments in cloud computing, artificial intelligence, and '
        'enterprise solutions continued to drive sustainable growth across all major business segments.'
    )
    doc.add_paragraph(
        'The company expanded its global workforce to 18,400 employees across 32 countries, '
        'strengthening our ability to serve clients in diverse markets. Net income reached $687 million, '
        'a 19.2% improvement year-over-year, reflecting both top-line growth and disciplined cost management.'
    )

    doc.add_page_break()

    # === Financial Highlights ===
    doc.add_heading('Financial Highlights', level=1)

    doc.add_heading('Revenue Analysis', level=2)
    doc.add_paragraph(
        'Total revenue for FY2025 was $4.82 billion, driven primarily by our cloud services division '
        'which grew 28% to $1.93 billion. Enterprise software licensing contributed $1.45 billion, '
        'while professional services accounted for $1.44 billion.'
    )

    doc.add_heading('Regional Revenue Breakdown', level=3)
    doc.add_paragraph(
        'North America remained our largest market at $2.41 billion (50%), followed by Europe at '
        '$1.30 billion (27%), Asia-Pacific at $867 million (18%), and Rest of World at $243 million (5%).'
    )

    doc.add_heading('Cost Management', level=2)
    doc.add_paragraph(
        'Operating expenses were carefully managed at $3.89 billion, achieving an operating margin '
        'of 19.3%, up from 17.8% in the prior year. Research and development spending increased to '
        '$723 million, representing 15% of revenue.'
    )

    doc.add_heading('Year-over-Year Comparison', level=3)
    doc.add_paragraph(
        'Compared to FY2024, gross margin improved by 1.2 percentage points to 62.4%, driven by '
        'favorable product mix shift toward higher-margin cloud subscriptions and reduced hardware costs.'
    )

    doc.add_heading('Balance Sheet Overview', level=2)
    doc.add_paragraph(
        'Total assets grew to $12.8 billion with cash and equivalents of $3.2 billion. Long-term debt '
        'was reduced by $450 million to $1.8 billion, strengthening our balance sheet flexibility.'
    )

    doc.add_page_break()

    # === Operations Review ===
    doc.add_heading('Operations Review', level=1)

    doc.add_heading('Product Development', level=2)
    doc.add_paragraph(
        'Our engineering teams launched 14 major product updates and 3 entirely new platform offerings '
        'during FY2025. The Meridian Cloud Platform 3.0 release in Q2 received industry recognition '
        'as a leader in the enterprise cloud infrastructure space.'
    )

    doc.add_heading('R&D Pipeline Status', level=3)
    doc.add_paragraph(
        'The R&D pipeline contains 47 active projects across four technology pillars: cloud infrastructure, '
        'AI/ML services, cybersecurity, and edge computing. Twelve projects are scheduled for release in H1 2026.'
    )

    doc.add_heading('Market Expansion', level=2)
    doc.add_paragraph(
        'Meridian expanded into 5 new markets during 2025, establishing offices in Singapore, '
        'Dubai, Sao Paulo, Seoul, and Warsaw. These expansion efforts are projected to generate '
        '$180 million in incremental revenue by FY2027.'
    )

    doc.add_heading('Asia-Pacific Expansion', level=3)
    doc.add_paragraph(
        'The Asia-Pacific region showed the strongest growth at 34% year-over-year, driven by '
        'enterprise cloud adoption in Japan, South Korea, and Southeast Asian markets. Our Singapore '
        'data center hub began operations in Q3 2025.'
    )

    doc.add_page_break()

    # === Strategic Initiatives ===
    doc.add_heading('Strategic Initiatives', level=1)

    doc.add_heading('Digital Transformation', level=2)
    doc.add_paragraph(
        'Our internal digital transformation program, codenamed Project Atlas, modernized core business '
        'processes across finance, HR, and supply chain operations. This initiative reduced operational '
        'overhead by $42 million annually while improving employee productivity metrics by 18%.'
    )

    doc.add_heading('Cloud Infrastructure Roadmap', level=3)
    doc.add_paragraph(
        'The three-year cloud infrastructure roadmap focuses on multi-cloud orchestration, serverless '
        'computing expansion, and AI-native development tools. Capital expenditure for cloud infrastructure '
        'is budgeted at $1.1 billion through FY2027.'
    )

    doc.add_heading('Sustainability Goals', level=2)
    doc.add_paragraph(
        'Meridian achieved a 32% reduction in carbon emissions from 2020 baseline levels, exceeding '
        'our intermediate target of 25%. All company-owned data centers now operate on 100% renewable '
        'energy, and we are on track to meet our net-zero commitment by 2030.'
    )

    doc.add_page_break()

    # === Corporate Governance ===
    doc.add_heading('Corporate Governance', level=1)
    doc.add_paragraph(
        'The Board of Directors maintained rigorous governance standards throughout FY2025. The board '
        'comprises 11 members, with 9 independent directors. Key actions during the year included the '
        'appointment of Dr. Elena Vasquez as Chief Technology Officer and the establishment of an AI '
        'Ethics Advisory Committee.'
    )
    doc.add_paragraph(
        'Shareholder returns remained a priority, with $1.2 billion returned through dividends and '
        'share repurchases. The quarterly dividend was increased by 8% to $0.54 per share, marking '
        'the 12th consecutive year of dividend increases.'
    )

    # Save document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
