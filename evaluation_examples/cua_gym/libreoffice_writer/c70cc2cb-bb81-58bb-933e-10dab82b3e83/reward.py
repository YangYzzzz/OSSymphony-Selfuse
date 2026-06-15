"""
Reward Script: Add fee schedule table in Section 4 of engagement letter
Task ID: writer_legal_064
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table exists with correct dimensions, placeholder removed
  Component 2 (0.30): Table data content matches expected fee schedule
  Component 3 (0.25): Header row bold text + gray background shading
  Component 4 (0.20): Hourly Rate column right-aligned in data rows
"""

import os

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_064'

# Expected table data (row 0 = header, rows 1-5 = data)
EXPECTED_HEADERS = ['Professional', 'Title', 'Hourly Rate']
EXPECTED_DATA = [
    ['J. Smith', 'Senior Partner', '$750'],
    ['A. Johnson', 'Junior Partner', '$550'],
    ['M. Williams', 'Senior Associate', '$400'],
    ['K. Brown', 'Associate', '$300'],
    ['L. Davis', 'Paralegal', '$175'],
]


def persist_app_state():
    """Save any unsaved LibreOffice state before verification."""
    try:
        os.environ["DISPLAY"] = ":0"
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        import time
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with correct dimensions AND placeholder removed (0.25 points)
    try:
        # Check placeholder is gone
        placeholder_gone = not any('[INSERT FEE SCHEDULE TABLE]' in p.text for p in doc.paragraphs)
        # Check table exists with 6 rows x 3 cols
        has_table = len(doc.tables) >= 1
        table_dims_ok = False
        if has_table:
            table = doc.tables[0]
            table_dims_ok = len(table.rows) == 6 and len(table.columns) == 3

        if placeholder_gone and has_table and table_dims_ok:
            print(f"PASS: Component 1 -- Table exists (6x3), placeholder removed (0.25 pts)")
            total_score += 0.25
        else:
            reasons = []
            if not placeholder_gone:
                reasons.append("placeholder still present")
            if not has_table:
                reasons.append("no table found")
            elif not table_dims_ok:
                reasons.append(f"table dims: {len(table.rows)}x{len(table.columns)}, expected 6x3")
            print(f"FAIL: Component 1 -- {'; '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Table data content matches (0.30 points)
    try:
        if len(doc.tables) < 1:
            print("FAIL: Component 2 -- No table to check data")
        else:
            table = doc.tables[0]
            # Check headers
            header_cells = [table.rows[0].cells[j].text.strip() for j in range(min(3, len(table.columns)))]
            headers_match = header_cells == EXPECTED_HEADERS

            # Check data rows
            data_match_count = 0
            for ri, expected_row in enumerate(EXPECTED_DATA):
                if ri + 1 < len(table.rows):
                    actual = [table.rows[ri + 1].cells[j].text.strip() for j in range(min(3, len(table.columns)))]
                    if actual == expected_row:
                        data_match_count += 1
                    else:
                        print(f"  Data mismatch row {ri+1}: expected {expected_row}, got {actual}")

            if headers_match and data_match_count == 5:
                print(f"PASS: Component 2 -- All table data matches (0.30 pts)")
                total_score += 0.30
            elif headers_match or data_match_count > 0:
                # Partial: award proportionally
                partial = 0.0
                if headers_match:
                    partial += 0.05
                partial += 0.05 * data_match_count
                partial = min(partial, 0.30)
                print(f"PARTIAL: Component 2 -- headers_match={headers_match}, data_rows={data_match_count}/5 ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 -- headers={header_cells}, data_match={data_match_count}/5")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Header row has bold text AND gray background shading (0.25 points)
    try:
        if len(doc.tables) < 1:
            print("FAIL: Component 3 -- No table to check formatting")
        else:
            table = doc.tables[0]
            bold_count = 0
            shading_count = 0

            for j in range(min(3, len(table.columns))):
                cell = table.rows[0].cells[j]
                # Check bold
                cell_bold = False
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.bold:
                            cell_bold = True
                            break
                    if cell_bold:
                        break
                if cell_bold:
                    bold_count += 1

                # Check shading (gray background)
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill_val = shd.get(qn('w:fill'))
                        if fill_val and fill_val.lower() not in ('auto', 'ffffff', '000000'):
                            shading_count += 1

            all_bold = bold_count == 3
            all_shaded = shading_count == 3

            if all_bold and all_shaded:
                print(f"PASS: Component 3 -- Header bold ({bold_count}/3) + gray shading ({shading_count}/3) (0.25 pts)")
                total_score += 0.25
            elif all_bold or all_shaded:
                partial = 0.125
                print(f"PARTIAL: Component 3 -- bold={bold_count}/3, shading={shading_count}/3 ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 -- bold={bold_count}/3, shading={shading_count}/3")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Hourly Rate column (col 2) right-aligned in data rows (0.20 points)
    try:
        if len(doc.tables) < 1:
            print("FAIL: Component 4 -- No table to check alignment")
        else:
            table = doc.tables[0]
            right_aligned_count = 0
            total_data_rows = 0

            for ri in range(1, len(table.rows)):
                total_data_rows += 1
                if len(table.rows[ri].cells) >= 3:
                    cell = table.rows[ri].cells[2]
                    for para in cell.paragraphs:
                        if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                            right_aligned_count += 1
                            break

            if total_data_rows > 0 and right_aligned_count == total_data_rows:
                print(f"PASS: Component 4 -- Rate column right-aligned ({right_aligned_count}/{total_data_rows}) (0.20 pts)")
                total_score += 0.20
            elif right_aligned_count > 0:
                partial = 0.20 * (right_aligned_count / max(total_data_rows, 1))
                print(f"PARTIAL: Component 4 -- right-aligned {right_aligned_count}/{total_data_rows} ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 -- right-aligned {right_aligned_count}/{total_data_rows}")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
