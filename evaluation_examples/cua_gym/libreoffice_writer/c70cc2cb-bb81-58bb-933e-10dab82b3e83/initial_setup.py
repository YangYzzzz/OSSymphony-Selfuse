"""
Initial Setup: Legal engagement letter with placeholder for fee schedule table
Task ID: writer_legal_064
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_064'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Firm letterhead ---
    header_para = doc.add_paragraph()
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header_para.add_run("HARRISON, BLAKE & ASSOCIATES LLP")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    addr_para = doc.add_paragraph()
    addr_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = addr_para.add_run("1200 Commerce Tower, Suite 4500\nChicago, Illinois 60601\nTel: (312) 555-8200 | Fax: (312) 555-8201")
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    # --- Separator line ---
    doc.add_paragraph("_" * 72)

    # --- Date and addressee ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(18)
    run = date_para.add_run("March 15, 2025")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    doc.add_paragraph()

    addr_lines = [
        "Ms. Catherine R. Whitfield",
        "Chief Executive Officer",
        "Meridian Healthcare Systems, Inc.",
        "8900 Lakeside Boulevard, 12th Floor",
        "Minneapolis, Minnesota 55402",
    ]
    for line in addr_lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(0)

    # --- Re line ---
    doc.add_paragraph()
    re_para = doc.add_paragraph()
    run = re_para.add_run("Re: Engagement Letter \u2014 Regulatory Compliance Review and Advisory Services")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    # --- Salutation ---
    doc.add_paragraph()
    sal = doc.add_paragraph()
    r = sal.add_run("Dear Ms. Whitfield:")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    # --- Section 1: Scope of Engagement ---
    s1_heading = doc.add_paragraph()
    r = s1_heading.add_run("1. Scope of Engagement")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    s1_body = doc.add_paragraph()
    r = s1_body.add_run(
        "This letter confirms the engagement of Harrison, Blake & Associates LLP "
        '("the Firm") by Meridian Healthcare Systems, Inc. ("the Client") to provide '
        "legal advisory services in connection with the Client's ongoing regulatory "
        "compliance obligations under federal and state healthcare regulations, "
        "including but not limited to the Health Insurance Portability and Accountability "
        "Act (HIPAA), the Affordable Care Act (ACA), and applicable state licensing requirements."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    s1_body.paragraph_format.space_after = Pt(6)

    s1_body2 = doc.add_paragraph()
    r = s1_body2.add_run(
        "The scope of our engagement shall include: (a) review and analysis of the Client's "
        "current compliance policies and procedures; (b) identification of regulatory gaps and "
        "risk areas; (c) preparation of remediation recommendations; (d) drafting of updated "
        "compliance documentation; and (e) ongoing advisory support for a period of twelve (12) months "
        "from the effective date of this agreement."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    # --- Section 2: Term and Termination ---
    s2_heading = doc.add_paragraph()
    s2_heading.paragraph_format.space_before = Pt(12)
    r = s2_heading.add_run("2. Term and Termination")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    s2_body = doc.add_paragraph()
    r = s2_body.add_run(
        "This engagement shall commence on April 1, 2025 and continue for a period of "
        "twelve (12) months, unless earlier terminated by either party upon thirty (30) days' "
        "written notice. In the event of termination, the Client shall remain responsible for "
        "payment of all fees and expenses incurred through the effective date of termination."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    # --- Section 3: Responsibilities ---
    s3_heading = doc.add_paragraph()
    s3_heading.paragraph_format.space_before = Pt(12)
    r = s3_heading.add_run("3. Client Responsibilities")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    s3_body = doc.add_paragraph()
    r = s3_body.add_run(
        "The Client agrees to provide timely access to all relevant documentation, personnel, "
        "and systems necessary for the Firm to perform its obligations under this engagement. "
        "The Client shall designate a primary point of contact who will be responsible for "
        "coordinating the flow of information between the Client and the Firm. The Client "
        "acknowledges that delays in providing requested materials may impact project timelines "
        "and deliverable schedules."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    # --- Section 4: Fees ---
    s4_heading = doc.add_paragraph()
    s4_heading.paragraph_format.space_before = Pt(12)
    r = s4_heading.add_run("4. Fees and Billing")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    s4_body = doc.add_paragraph()
    r = s4_body.add_run(
        "The Firm shall bill the Client on a monthly basis for professional services rendered. "
        "All invoices are due and payable within thirty (30) days of receipt. The applicable "
        "hourly rates for personnel assigned to this engagement are set forth below:"
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    s4_body.paragraph_format.space_after = Pt(6)

    # --- Placeholder for fee schedule table ---
    placeholder = doc.add_paragraph()
    r = placeholder.add_run("[INSERT FEE SCHEDULE TABLE]")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    s4_body2 = doc.add_paragraph()
    s4_body2.paragraph_format.space_before = Pt(6)
    r = s4_body2.add_run(
        "In addition to professional fees, the Client shall reimburse the Firm for all "
        "reasonable out-of-pocket expenses incurred in connection with this engagement, "
        "including but not limited to travel, filing fees, and third-party vendor costs. "
        "Expenses exceeding $500 individually shall require prior written approval from the Client."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    # --- Section 5: Confidentiality ---
    s5_heading = doc.add_paragraph()
    s5_heading.paragraph_format.space_before = Pt(12)
    r = s5_heading.add_run("5. Confidentiality")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    s5_body = doc.add_paragraph()
    r = s5_body.add_run(
        "The Firm shall maintain the confidentiality of all information provided by the Client "
        "in connection with this engagement, in accordance with applicable rules of professional "
        "conduct and attorney-client privilege. This obligation shall survive the termination "
        "of this engagement."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    # --- Section 6: Limitation of Liability ---
    s6_heading = doc.add_paragraph()
    s6_heading.paragraph_format.space_before = Pt(12)
    r = s6_heading.add_run("6. Limitation of Liability")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    s6_body = doc.add_paragraph()
    r = s6_body.add_run(
        "The Firm's total aggregate liability arising out of or relating to this engagement "
        "shall not exceed the total fees paid by the Client to the Firm during the twelve-month "
        "period preceding the claim. In no event shall the Firm be liable for any indirect, "
        "consequential, special, or punitive damages."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    # --- Closing ---
    doc.add_paragraph()
    closing = doc.add_paragraph()
    r = closing.add_run(
        "We appreciate the opportunity to serve Meridian Healthcare Systems and look forward "
        "to a productive engagement. Please indicate your acceptance of the terms set forth "
        "herein by signing below and returning a copy of this letter to our office."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    doc.add_paragraph()
    sincerely = doc.add_paragraph()
    r = sincerely.add_run("Very truly yours,")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_paragraph()
    r = sig.add_run("James R. Harrison")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.bold = True

    title_para = doc.add_paragraph()
    r = title_para.add_run("Managing Partner")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    firm = doc.add_paragraph()
    r = firm.add_run("Harrison, Blake & Associates LLP")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
