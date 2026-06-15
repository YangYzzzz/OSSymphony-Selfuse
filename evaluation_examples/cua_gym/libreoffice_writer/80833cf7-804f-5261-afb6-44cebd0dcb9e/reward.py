"""
Reward Script: Apply Academic_Report template styles to document
Task ID: writer_bs_065
Domain: libreoffice_writer
Scoring:
  C1 (0.25) - Heading 1 style has template formatting (Georgia, larger size, new color)
  C2 (0.25) - Heading 2 style has template formatting (Georgia, italic, new color/size)
  C3 (0.20) - Body Text style has template formatting (Palatino Linotype, color, line spacing)
  C4 (0.15) - Caption style has template formatting (Georgia, italic, centered, new color)
  C5 (0.15) - New custom styles (Academic Note, Academic Subtitle) added to style catalog
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_065'


def persist_app_state(domain):
    """Save any unsaved edits in LibreOffice."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        import time
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def color_distance(c1, c2):
    """Simple RGB distance."""
    from math import sqrt
    return sqrt(sum((a - b) ** 2 for a, b in zip(
        (c1[0], c1[1], c1[2]),
        (c2[0], c2[1], c2[2])
    )))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: text content must be preserved (121 paragraphs, 25000+ chars)
    all_text = ''.join(p.text for p in doc.paragraphs)
    if len(doc.paragraphs) < 100 or len(all_text) < 20000:
        print(f"PRECONDITION FAIL: Document content appears damaged. "
              f"Paragraphs={len(doc.paragraphs)}, chars={len(all_text)}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Heading 1 style has template formatting (0.25 points)
    # Golden: font.name=Georgia, font.size=304800 (24pt), color=1A3C6E
    # Initial: font.name=None, font.size=177800 (14pt), color=365F91
    try:
        h1_style = doc.styles['Heading 1']
        h1_font = h1_style.font
        checks_passed = 0
        total_checks = 3

        # Check font name changed to Georgia
        if h1_font.name and h1_font.name.lower() == 'georgia':
            checks_passed += 1
            print(f"  PASS: Heading 1 font.name = {h1_font.name}")
        else:
            print(f"  FAIL: Heading 1 font.name = {h1_font.name}, expected Georgia")

        # Check font size changed to 304800 (24pt)
        if h1_font.size and h1_font.size > 200000:
            checks_passed += 1
            print(f"  PASS: Heading 1 font.size = {h1_font.size} (>200000)")
        else:
            print(f"  FAIL: Heading 1 font.size = {h1_font.size}, expected >200000")

        # Check color changed to 1A3C6E (different from initial 365F91)
        if h1_font.color and h1_font.color.rgb:
            rgb = h1_font.color.rgb
            # Check it's close to 1A3C6E (dark navy) - differs from initial 365F91
            dist_to_golden = color_distance(rgb, RGBColor(0x1A, 0x3C, 0x6E))
            dist_to_initial = color_distance(rgb, RGBColor(0x36, 0x5F, 0x91))
            if dist_to_golden < dist_to_initial:
                checks_passed += 1
                print(f"  PASS: Heading 1 color = {rgb} (closer to golden 1A3C6E)")
            else:
                print(f"  FAIL: Heading 1 color = {rgb} (still closer to initial 365F91)")
        else:
            print(f"  FAIL: Heading 1 color = None")

        if checks_passed >= 2:
            earned = 0.25
            total_score += earned
            print(f"PASS: Component 1 — Heading 1 template formatting ({checks_passed}/{total_checks} checks) ({earned} pts)")
        else:
            print(f"FAIL: Component 1 — Heading 1 template formatting ({checks_passed}/{total_checks} checks)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Heading 2 style has template formatting (0.25 points)
    # Golden: font.name=Georgia, font.size=228600 (18pt), bold=True, italic=True, color=2E5C8A
    # Initial: font.name=None, font.size=165100 (13pt), bold=True, italic=None, color=4F81BD
    try:
        h2_style = doc.styles['Heading 2']
        h2_font = h2_style.font
        checks_passed = 0
        total_checks = 3

        # Check font name changed to Georgia
        if h2_font.name and h2_font.name.lower() == 'georgia':
            checks_passed += 1
            print(f"  PASS: Heading 2 font.name = {h2_font.name}")
        else:
            print(f"  FAIL: Heading 2 font.name = {h2_font.name}, expected Georgia")

        # Check italic is now True (was None in initial)
        if h2_font.italic is True:
            checks_passed += 1
            print(f"  PASS: Heading 2 italic = True")
        else:
            print(f"  FAIL: Heading 2 italic = {h2_font.italic}, expected True")

        # Check color changed to 2E5C8A (differs from initial 4F81BD)
        if h2_font.color and h2_font.color.rgb:
            rgb = h2_font.color.rgb
            dist_to_golden = color_distance(rgb, RGBColor(0x2E, 0x5C, 0x8A))
            dist_to_initial = color_distance(rgb, RGBColor(0x4F, 0x81, 0xBD))
            if dist_to_golden < dist_to_initial:
                checks_passed += 1
                print(f"  PASS: Heading 2 color = {rgb} (closer to golden 2E5C8A)")
            else:
                print(f"  FAIL: Heading 2 color = {rgb} (still closer to initial 4F81BD)")
        else:
            print(f"  FAIL: Heading 2 color = None")

        if checks_passed >= 2:
            earned = 0.25
            total_score += earned
            print(f"PASS: Component 2 — Heading 2 template formatting ({checks_passed}/{total_checks} checks) ({earned} pts)")
        else:
            print(f"FAIL: Component 2 — Heading 2 template formatting ({checks_passed}/{total_checks} checks)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Body Text style has template formatting (0.20 points)
    # Golden: font.name=Palatino Linotype, font.size=139700 (11pt), color=333333, line_spacing=1.15
    # Initial: font.name=None, font.size=None, color=None, line_spacing=None
    try:
        bt_style = doc.styles['Body Text']
        bt_font = bt_style.font
        bt_pf = bt_style.paragraph_format
        checks_passed = 0
        total_checks = 3

        # Check font name changed to Palatino Linotype
        if bt_font.name and 'palatino' in bt_font.name.lower():
            checks_passed += 1
            print(f"  PASS: Body Text font.name = {bt_font.name}")
        else:
            print(f"  FAIL: Body Text font.name = {bt_font.name}, expected Palatino Linotype")

        # Check color changed from None to 333333
        if bt_font.color and bt_font.color.rgb:
            rgb = bt_font.color.rgb
            dist = color_distance(rgb, RGBColor(0x33, 0x33, 0x33))
            if dist < 50:
                checks_passed += 1
                print(f"  PASS: Body Text color = {rgb} (close to 333333)")
            else:
                print(f"  FAIL: Body Text color = {rgb}, expected close to 333333")
        else:
            print(f"  FAIL: Body Text color = None, expected 333333")

        # Check line spacing changed from None to ~1.15
        if bt_pf.line_spacing is not None:
            ls = float(bt_pf.line_spacing)
            if 1.1 <= ls <= 1.2:
                checks_passed += 1
                print(f"  PASS: Body Text line_spacing = {ls}")
            else:
                print(f"  FAIL: Body Text line_spacing = {ls}, expected ~1.15")
        else:
            print(f"  FAIL: Body Text line_spacing = None, expected ~1.15")

        if checks_passed >= 2:
            earned = 0.20
            total_score += earned
            print(f"PASS: Component 3 — Body Text template formatting ({checks_passed}/{total_checks} checks) ({earned} pts)")
        else:
            print(f"FAIL: Component 3 — Body Text template formatting ({checks_passed}/{total_checks} checks)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Caption style has template formatting (0.15 points)
    # Golden: font.name=Georgia, italic=True, color=666666, alignment=CENTER, space_before=50800
    # Initial: font.name=None, italic=None, color=4F81BD, alignment=None, space_before=None
    try:
        cap_style = doc.styles['Caption']
        cap_font = cap_style.font
        cap_pf = cap_style.paragraph_format
        checks_passed = 0
        total_checks = 3

        # Check italic changed from None to True
        if cap_font.italic is True:
            checks_passed += 1
            print(f"  PASS: Caption italic = True")
        else:
            print(f"  FAIL: Caption italic = {cap_font.italic}, expected True")

        # Check color changed from 4F81BD to 666666
        if cap_font.color and cap_font.color.rgb:
            rgb = cap_font.color.rgb
            dist_to_golden = color_distance(rgb, RGBColor(0x66, 0x66, 0x66))
            dist_to_initial = color_distance(rgb, RGBColor(0x4F, 0x81, 0xBD))
            if dist_to_golden < dist_to_initial:
                checks_passed += 1
                print(f"  PASS: Caption color = {rgb} (closer to golden 666666)")
            else:
                print(f"  FAIL: Caption color = {rgb} (still closer to initial 4F81BD)")
        else:
            print(f"  FAIL: Caption color = None")

        # Check alignment changed from None to CENTER
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        if cap_pf.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            checks_passed += 1
            print(f"  PASS: Caption alignment = CENTER")
        else:
            print(f"  FAIL: Caption alignment = {cap_pf.alignment}, expected CENTER")

        if checks_passed >= 2:
            earned = 0.15
            total_score += earned
            print(f"PASS: Component 4 — Caption template formatting ({checks_passed}/{total_checks} checks) ({earned} pts)")
        else:
            print(f"FAIL: Component 4 — Caption template formatting ({checks_passed}/{total_checks} checks)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: New custom styles added from template (0.15 points)
    # Golden has: Academic Note, Academic Subtitle (built_in=False)
    # Initial does NOT have these styles
    try:
        custom_styles_found = 0
        expected_customs = ['Academic Note', 'Academic Subtitle']

        for sname in expected_customs:
            try:
                s = doc.styles[sname]
                if not s.builtin:
                    custom_styles_found += 1
                    print(f"  PASS: Custom style '{sname}' found (builtin={s.builtin})")
                else:
                    print(f"  FAIL: Style '{sname}' is builtin, expected custom")
            except KeyError:
                print(f"  FAIL: Custom style '{sname}' not found")

        if custom_styles_found >= 2:
            earned = 0.15
            total_score += earned
            print(f"PASS: Component 5 — New custom styles from template ({custom_styles_found}/{len(expected_customs)}) ({earned} pts)")
        elif custom_styles_found >= 1:
            earned = 0.075
            total_score += earned
            print(f"PARTIAL: Component 5 — {custom_styles_found}/{len(expected_customs)} custom styles ({earned} pts)")
        else:
            print(f"FAIL: Component 5 — No custom styles from template found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
