"""
Initial Setup: Apply Academic_Report template styles to a Writer document
Task ID: writer_bs_065
Domain: libreoffice_writer

Creates:
1. A 20-page Writer document with default styles (/home/user/writer_bs_065.docx)
2. An Academic_Report.ott template with custom Heading, Body Text, and Caption styles
3. Opens the document in LibreOffice Writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_065'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
TEMPLATE_DOCX = f'{WORKDIR}/Academic_Report_temp.docx'
TEMPLATE_OTT = f'{WORKDIR}/Academic_Report.ott'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_main_document():
    """Create a 20-page document with default styles and realistic content."""
    doc = Document()

    # Set default page margins for consistent page count
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # --- Page 1: Title page ---
    doc.add_heading('Comprehensive Analysis of Renewable Energy Adoption in Urban Centers', level=1)
    doc.add_paragraph('')
    p = doc.add_paragraph('Prepared by the Institute for Sustainable Development')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph('Research Division — Environmental Policy Group')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph('March 2025')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('')
    doc.add_paragraph('This report presents findings from a two-year longitudinal study examining the adoption patterns of renewable energy technologies across 45 major urban centers in North America, Europe, and Asia-Pacific regions. The study encompasses solar photovoltaic installations, wind energy capacity, geothermal systems, and emerging hydrogen fuel cell deployments.')

    doc.add_page_break()

    # --- Page 2-3: Table of Contents / Executive Summary ---
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph('The global transition to renewable energy sources has accelerated significantly over the past decade, driven by declining technology costs, supportive policy frameworks, and growing public awareness of climate change impacts. This comprehensive analysis reveals that urban centers with integrated policy approaches achieve 2.3 times higher adoption rates compared to cities relying solely on market-driven mechanisms.')
    doc.add_paragraph('Key findings from our analysis include a 47% increase in residential solar installations across surveyed cities between 2023 and 2025, with the highest growth rates observed in cities offering streamlined permitting processes and community solar programs. Commercial and industrial sectors showed a 31% increase in on-site renewable generation capacity, primarily driven by corporate sustainability commitments and favorable power purchase agreement structures.')
    doc.add_paragraph('The economic analysis demonstrates that cities with robust renewable energy ecosystems experienced 12% higher job creation in the clean energy sector compared to the national average. Furthermore, property values in neighborhoods with high solar penetration rates showed a 3.8% premium over comparable areas without significant renewable installations.')
    doc.add_paragraph('Our recommendations emphasize the importance of coordinated policy design, workforce development programs, equitable access to clean energy benefits, and infrastructure modernization to support bidirectional energy flows. The report concludes with a framework for cities at different stages of their renewable energy transition journey.')

    doc.add_page_break()

    # --- Pages 3-5: Chapter 1 ---
    doc.add_heading('Chapter 1: Introduction and Background', level=1)
    doc.add_heading('1.1 Research Context', level=2)
    doc.add_paragraph('The imperative to transition toward sustainable energy systems has never been more pressing. Global carbon dioxide emissions reached 36.8 billion tonnes in 2024, with urban areas accounting for approximately 70% of energy-related emissions. Cities, as both primary contributors to greenhouse gas emissions and frontline responders to climate impacts, occupy a critical position in the global decarbonization effort.')
    doc.add_paragraph('Previous research has established the technical feasibility of achieving high renewable energy penetration in urban settings. However, the socioeconomic, regulatory, and behavioral factors influencing adoption rates remain incompletely understood. This study addresses this gap by combining quantitative deployment data with qualitative policy analysis across a diverse sample of urban centers.')

    doc.add_heading('1.2 Study Objectives', level=2)
    doc.add_paragraph('The primary objectives of this research are threefold. First, we seek to quantify current renewable energy adoption rates across different urban typologies, including compact high-density cities, sprawling metropolitan regions, and mid-sized urban centers. Second, we analyze the policy instruments and market conditions most strongly correlated with accelerated adoption. Third, we develop predictive models for future deployment trajectories under various policy scenarios.')
    doc.add_paragraph('Caption: Figure 1.1 — Global renewable energy capacity growth, 2015-2025')
    doc.add_paragraph('The analytical framework integrates data from municipal energy databases, utility interconnection records, building permit systems, and satellite imagery analysis. Machine learning algorithms were employed to identify spatial patterns in installation density and predict areas with high untapped potential for renewable energy deployment.')

    doc.add_heading('1.3 Methodology Overview', level=2)
    doc.add_paragraph('Our mixed-methods approach combines quantitative data analysis with semi-structured interviews of key stakeholders in each surveyed city. The quantitative component draws on installation-level data for over 2.3 million renewable energy systems, encompassing residential rooftop solar, commercial-scale solar arrays, small and medium wind turbines, ground-source heat pumps, and pilot hydrogen fuel cell installations.')
    doc.add_paragraph('Statistical analysis employed hierarchical linear modeling to account for the nested structure of installations within neighborhoods, neighborhoods within cities, and cities within regional regulatory frameworks. Spatial econometric techniques, including geographically weighted regression, captured the variation in adoption drivers across different geographic contexts.')

    doc.add_page_break()

    # --- Pages 5-8: Chapter 2 ---
    doc.add_heading('Chapter 2: Solar Energy Deployment Analysis', level=1)
    doc.add_heading('2.1 Residential Solar Installations', level=2)
    doc.add_paragraph('Residential solar photovoltaic systems represent the fastest-growing segment of urban renewable energy deployment. Our analysis of permitting data reveals that the median residential system size increased from 6.2 kW in 2022 to 8.7 kW in 2025, reflecting both declining per-watt costs and growing homeowner appetite for energy independence.')
    doc.add_paragraph('The geographic distribution of residential solar installations follows distinct patterns correlated with building stock characteristics, income demographics, and local policy environments. Cities in the Sun Belt region of the United States continue to lead in absolute installation numbers, with Phoenix, Las Vegas, and San Antonio recording the highest per-capita deployment rates.')
    doc.add_paragraph('However, the most dramatic growth rates were observed in cities that recently implemented supportive policies. Minneapolis experienced a 340% increase in residential solar permits following the introduction of its Solar*Rewards community program. Similarly, Portland\'s Solarize Portland initiative drove a 280% increase in installations in participating neighborhoods.')

    doc.add_heading('2.2 Commercial and Industrial Solar', level=2)
    doc.add_paragraph('The commercial and industrial solar segment showed robust growth, with total installed capacity increasing by 31% across our sample cities. Large retail chains, data center operators, and manufacturing facilities emerged as the primary drivers of this growth, motivated by a combination of cost savings, sustainability reporting requirements, and customer expectations.')
    doc.add_paragraph('Caption: Table 2.1 — Top 10 cities by commercial solar capacity additions, 2024-2025')
    doc.add_paragraph('On-site solar generation at commercial facilities reached an average capacity factor of 18.3% across our sample, with significant variation between regions. Desert Southwest installations achieved capacity factors exceeding 25%, while Pacific Northwest sites averaged 14.7%. These differences underscore the importance of location-specific financial modeling in investment decisions.')
    doc.add_paragraph('The emergence of solar carport structures has been a notable trend in the commercial segment. Parking lot solar installations accounted for 23% of new commercial solar capacity in 2025, up from just 8% in 2022. These dual-use installations provide both renewable energy generation and shaded parking, creating additional value propositions for commercial property owners.')

    doc.add_heading('2.3 Community Solar Programs', level=2)
    doc.add_paragraph('Community solar programs have emerged as a critical mechanism for extending solar access to households unable to host rooftop systems. Approximately 35% of residential electricity customers in urban areas face barriers to on-site solar installation, including rental housing, shaded rooftops, structural limitations, or insufficient credit scores for financing.')
    doc.add_paragraph('Our analysis identified 127 active community solar programs across the surveyed cities, with a total subscribed capacity of 4.8 GW. Programs structured as utility-administered models showed the highest participation rates among low-to-moderate income households, while third-party developer models achieved lower customer acquisition costs and faster deployment timelines.')

    doc.add_page_break()

    # --- Pages 8-11: Chapter 3 ---
    doc.add_heading('Chapter 3: Wind Energy in Urban Environments', level=1)
    doc.add_heading('3.1 Small Wind Turbine Deployments', level=2)
    doc.add_paragraph('Urban wind energy presents unique challenges and opportunities compared to utility-scale rural installations. Building-mounted and small freestanding wind turbines have seen modest but growing adoption in cities with favorable wind resources. Our data encompasses 3,847 small wind installations across the surveyed cities, with a combined capacity of 28.4 MW.')
    doc.add_paragraph('The performance of urban wind installations varies significantly based on turbine placement, local terrain effects, and building wake interactions. Computational fluid dynamics modeling conducted for 12 cities in our sample revealed that rooftop-mounted turbines at heights exceeding 1.5 times the surrounding building height achieved capacity factors 40% higher than turbines positioned at building height.')
    doc.add_paragraph('Caption: Figure 3.1 — Wind speed profiles across urban morphology types')
    doc.add_paragraph('Noise considerations remain a significant barrier to urban wind deployment. Despite advances in blade design that have reduced acoustic emissions by an average of 8 dB over the past five years, community resistance to wind turbine noise continues to limit installations in residential areas. Cities that established clear noise ordinance thresholds specifically calibrated for small wind systems reported 2.1 times higher installation rates than those applying general industrial noise standards.')

    doc.add_heading('3.2 Urban Wind Resource Assessment', level=2)
    doc.add_paragraph('Accurate wind resource assessment in complex urban terrain requires sophisticated modeling techniques that account for building-induced turbulence, thermal convection patterns, and seasonal variation in atmospheric stability. Traditional wind resource maps developed for rural areas significantly overestimate available energy at typical urban installation heights.')
    doc.add_paragraph('Our collaboration with meteorological research institutions produced high-resolution wind atlases for 15 cities, incorporating LiDAR measurements, weather station data, and building geometry databases. Results indicate that average annual wind speeds at 30-meter height above ground range from 3.2 m/s in sheltered downtown canyons to 6.8 m/s at exposed coastal and hilltop locations within the same city.')
    doc.add_paragraph('The economic viability threshold for small wind turbines, typically cited at 4.5 m/s average annual wind speed, is met at suitable locations in 28 of the 45 surveyed cities. However, when accounting for urban turbulence intensity and its impact on turbine fatigue loading, the effective economic threshold increases to approximately 5.2 m/s, reducing the number of viable cities to 19.')

    doc.add_heading('3.3 Offshore Wind Contributions to Urban Energy Supply', level=2)
    doc.add_paragraph('While not physically located within urban boundaries, offshore wind projects increasingly contribute to urban renewable energy portfolios through power purchase agreements and green tariff programs. Coastal cities in our sample reported that offshore wind commitments account for an average of 15% of their renewable energy procurement targets for 2030.')
    doc.add_paragraph('The development pipeline for offshore wind projects serving urban markets has grown substantially, with 47 GW of capacity in various stages of development across the Atlantic seaboard, Great Lakes region, and Pacific Coast. Supply chain investments associated with these projects are creating significant economic development opportunities in port cities, with an estimated 12,400 direct manufacturing and installation jobs expected by 2028.')

    doc.add_page_break()

    # --- Pages 11-14: Chapter 4 ---
    doc.add_heading('Chapter 4: Geothermal and Ground-Source Heat Pump Systems', level=1)
    doc.add_heading('4.1 Ground-Source Heat Pump Adoption', level=2)
    doc.add_paragraph('Ground-source heat pump (GSHP) systems represent an often-overlooked component of urban renewable energy strategies. By leveraging the stable temperatures found 6 to 300 feet below the earth\'s surface, these systems can provide heating and cooling at efficiencies 3 to 5 times greater than conventional HVAC equipment. Our survey identified 14,230 GSHP installations across the 45 cities, with a combined thermal capacity of 892 MW.')
    doc.add_paragraph('The adoption of GSHP technology shows strong correlation with heating-dominated climates, high natural gas prices, and the availability of qualified installation contractors. Cities in the U.S. Northeast and Midwest, Scandinavia, and northern China recorded the highest installation densities. Stockholm leads globally with over 200,000 residential GSHP systems serving a metropolitan population of 2.4 million.')
    doc.add_paragraph('New construction projects incorporating GSHP systems from the design phase achieve significantly lower installation costs compared to retrofit projects. Our analysis shows that GSHP integration in new commercial buildings adds 2-5% to total construction costs but reduces annual HVAC energy consumption by 45-65%, yielding simple payback periods of 5 to 8 years in most climate zones.')

    doc.add_heading('4.2 District Geothermal Networks', level=2)
    doc.add_paragraph('District geothermal networks, also known as geo-exchange or ambient loop systems, represent an emerging approach to community-scale thermal energy management. These networks circulate water through shared underground loop fields to provide heating and cooling to multiple buildings, achieving higher system-level efficiencies through load diversity and thermal energy sharing between buildings with complementary heating and cooling demands.')
    doc.add_paragraph('Caption: Figure 4.1 — Schematic of a fifth-generation district thermal network')
    doc.add_paragraph('Our research documented 34 operational or planned district geothermal networks across the surveyed cities. The largest operational system serves a 45-hectare mixed-use development in Toronto, providing heating and cooling to 8,500 residential units and 120,000 square meters of commercial space. System performance monitoring shows a seasonal coefficient of performance averaging 4.2 for heating and 6.1 for cooling.')
    doc.add_paragraph('The economics of district geothermal networks improve substantially with scale and building density. Our modeling indicates that networks serving developments with floor area ratios above 2.0 achieve levelized costs of thermal energy 15-30% below conventional alternatives when evaluated over a 30-year system lifetime, including the significant value of avoided carbon emissions.')

    doc.add_heading('4.3 Enhanced Geothermal Systems', level=2)
    doc.add_paragraph('Enhanced geothermal systems (EGS) technology, which creates artificial geothermal reservoirs by hydraulic stimulation of hot dry rock formations, has advanced from laboratory demonstrations to pilot-scale deployments. While current EGS projects are located in rural areas, the baseload electricity generation capability of these systems holds significant potential for urban energy supply.')
    doc.add_paragraph('Three cities in our sample — Boise, Idaho; Reykjavik, Iceland; and Munich, Germany — already derive a substantial portion of their heating energy from geothermal resources. Boise\'s geothermal district heating system, one of the largest in the United States, serves 92 buildings in the downtown core using naturally occurring hot water at 77 degrees Celsius.')

    doc.add_page_break()

    # --- Pages 14-17: Chapter 5 ---
    doc.add_heading('Chapter 5: Policy Framework Analysis', level=1)
    doc.add_heading('5.1 Financial Incentive Structures', level=2)
    doc.add_paragraph('The design of financial incentive programs significantly influences renewable energy adoption rates. Our cross-city analysis evaluated five categories of financial incentives: direct rebates, tax credits, performance-based incentives, low-interest loan programs, and property-assessed clean energy (PACE) financing.')
    doc.add_paragraph('Direct rebates, while effective at stimulating initial adoption, show diminishing returns as market maturity increases. Cities that transitioned from upfront rebates to performance-based incentives reported sustained deployment growth rates, as the performance-based structure attracted higher-quality installations and more sophisticated market participants.')
    doc.add_paragraph('Caption: Table 5.1 — Incentive program comparison across 15 leading cities')
    doc.add_paragraph('PACE financing programs demonstrated the strongest correlation with adoption among low-to-moderate income households, as they eliminate the upfront cost barrier while allowing repayment through property tax assessments. However, concerns about consumer protection and the priority lien status of PACE assessments have limited program availability in some jurisdictions.')

    doc.add_heading('5.2 Regulatory and Permitting Reforms', level=2)
    doc.add_paragraph('Permitting process efficiency emerged as one of the strongest predictors of residential solar adoption rates. Cities with streamlined permitting processes — defined as online application submission, same-day plan review for standard systems, and automated interconnection approval — achieved installation rates 3.2 times higher than cities with traditional paper-based, multi-week permitting workflows.')
    doc.add_paragraph('The SolarAPP+ automated permitting platform, deployed in 37 of our surveyed cities, reduced average residential solar permit review times from 15.3 business days to less than one day. Installer surveys indicated that the reduction in soft costs associated with permitting delays contributed to a 7-12% decrease in total system prices in participating markets.')
    doc.add_paragraph('Building code requirements for solar readiness in new construction have been adopted by 22 cities in our sample. These codes typically require pre-wiring for solar systems, structural reinforcement for rooftop mounting, and reserved electrical panel capacity. Compliance cost analyses show that solar-ready construction adds less than 0.3% to total building costs while reducing future solar installation costs by 15-25%.')

    doc.add_heading('5.3 Grid Integration and Net Metering Policies', level=2)
    doc.add_paragraph('Net metering policies and their successors continue to shape the economic proposition for distributed renewable energy. Our analysis found that cities served by utilities offering full retail rate net metering experienced 2.8 times higher residential solar adoption compared to cities with reduced compensation rates or buy-all, sell-all rate structures.')
    doc.add_paragraph('The transition from traditional net metering to time-of-use rates and net billing arrangements is reshaping the value proposition for residential solar. In cities where evening peak rates exceed midday rates by a factor of two or more, solar-plus-storage systems have achieved market penetration rates 4.5 times higher than in cities with flat rate structures.')
    doc.add_paragraph('Grid capacity constraints are emerging as a significant bottleneck for continued renewable energy deployment in several urban areas. Distribution system upgrade costs, interconnection queue backlogs, and hosting capacity limitations affected 18 of our 45 surveyed cities. Utilities that proactively invested in grid modernization and published transparent hosting capacity maps reported fewer interconnection delays and higher customer satisfaction.')

    doc.add_page_break()

    # --- Pages 17-19: Chapter 6 ---
    doc.add_heading('Chapter 6: Economic Impact Assessment', level=1)
    doc.add_heading('6.1 Employment Effects', level=2)
    doc.add_paragraph('The renewable energy sector has become a significant source of employment in urban economies. Our analysis of labor market data reveals that the 45 surveyed cities collectively support approximately 487,000 clean energy jobs, encompassing solar installation, wind turbine manufacturing and maintenance, energy efficiency services, grid modernization, and electric vehicle infrastructure development.')
    doc.add_paragraph('Solar installation represents the largest single job category, accounting for 38% of clean energy employment in our sample. The median hourly wage for solar installers across surveyed cities was $24.50, with significant variation between markets — ranging from $18.75 in lower-cost metropolitan areas to $38.20 in high-cost coastal cities.')
    doc.add_paragraph('Workforce development programs specifically targeting clean energy careers have been established in 31 of the surveyed cities. Programs offering a combination of classroom instruction, hands-on training, and paid apprenticeships demonstrated the highest completion and job placement rates, with 78% of graduates securing employment in the clean energy sector within six months.')

    doc.add_heading('6.2 Property Value Impacts', level=2)
    doc.add_paragraph('The relationship between renewable energy installations and property values has important implications for housing equity and tax revenue. Our hedonic pricing analysis, controlling for building characteristics, neighborhood amenities, and market conditions, found that homes with owned solar systems sold at a premium of 3.8% compared to comparable properties without solar.')
    doc.add_paragraph('The property value premium associated with solar installations varied by market conditions and system characteristics. Newer systems with battery storage commanded premiums of 5.2%, while older systems nearing the end of their warranty periods showed premiums of only 1.4%. Leased solar systems showed no statistically significant impact on sale prices.')
    doc.add_paragraph('Caption: Figure 6.1 — Solar premium by housing market segment and system age')
    doc.add_paragraph('Neighborhood-level effects were also observed. Census tracts with solar penetration rates exceeding 15% showed aggregate property value increases of 1.2% beyond what individual installation premiums would predict, suggesting positive spillover effects from visible renewable energy adoption on neighborhood desirability and perceived environmental quality.')

    doc.add_heading('6.3 Municipal Revenue and Energy Cost Savings', level=2)
    doc.add_paragraph('Municipal governments in our sample reported combined annual energy cost savings of $1.34 billion from renewable energy installations on government-owned buildings and facilities. School districts accounted for the largest share of public sector solar deployment, with 2,847 schools across the 45 cities hosting solar installations with a combined capacity of 1.2 GW.')
    doc.add_paragraph('The fiscal impact of renewable energy deployment extends beyond direct energy savings. Sales tax revenue from solar equipment purchases, increased property tax collections from value-enhanced homes, and business tax revenue from clean energy companies contributed an estimated $890 million annually to municipal budgets across our sample cities.')

    doc.add_page_break()

    # --- Pages 19-20: Chapter 7 and Conclusions ---
    doc.add_heading('Chapter 7: Conclusions and Recommendations', level=1)
    doc.add_heading('7.1 Summary of Key Findings', level=2)
    doc.add_paragraph('This comprehensive analysis demonstrates that urban renewable energy adoption has reached a tipping point in many cities, transitioning from niche environmental initiatives to mainstream infrastructure investments. The convergence of declining technology costs, supportive policy environments, and growing demand from residents and businesses is creating self-reinforcing market dynamics that accelerate deployment.')
    doc.add_paragraph('The most successful cities share several common characteristics: coordinated policy frameworks that address multiple barriers simultaneously, proactive grid infrastructure investment, workforce development programs aligned with industry needs, and intentional equity provisions ensuring broad access to clean energy benefits.')
    doc.add_paragraph('However, significant challenges remain. Grid integration constraints, permitting bottlenecks, workforce shortages in specialized trades, and persistent inequities in clean energy access threaten to slow progress unless deliberately addressed. The transition also raises important questions about energy system governance, utility business model evolution, and the appropriate allocation of infrastructure costs.')

    doc.add_heading('7.2 Recommendations for City Leaders', level=2)
    doc.add_paragraph('Based on our analysis, we offer the following recommendations for city leaders seeking to accelerate renewable energy adoption:')
    doc.add_paragraph('First, adopt comprehensive clean energy plans with legally binding targets and transparent progress tracking mechanisms. Cities with codified targets achieved deployment rates 1.8 times higher than those with aspirational goals alone.')
    doc.add_paragraph('Second, streamline permitting and interconnection processes through digitization and automation. Every additional week of permitting delay reduces residential solar adoption rates by approximately 4%.')
    doc.add_paragraph('Third, invest in grid modernization proactively, guided by transparent hosting capacity analysis and long-range distribution system planning. Reactive grid upgrades are significantly more expensive and create deployment bottlenecks.')
    doc.add_paragraph('Fourth, design financial incentive programs that evolve with market maturity, transitioning from upfront rebates to performance-based incentives and innovative financing mechanisms as markets develop.')
    doc.add_paragraph('Fifth, ensure equitable access to clean energy benefits through targeted programs serving low-to-moderate income households, renters, and historically underserved communities.')
    doc.add_paragraph('Caption: Table 7.1 — Recommended policy priority matrix by city readiness level')

    doc.add_heading('References', level=1)
    doc.add_paragraph('Anderson, T. K., & Williams, R. J. (2024). "Urban solar adoption dynamics: A spatial panel analysis." Energy Policy, 178, 113542.')
    doc.add_paragraph('Chen, M., Patel, S., & Okafor, E. (2024). "Community solar programs and energy equity." Renewable Energy, 215, 118940.')
    doc.add_paragraph('Davis, K. L., & Martinez, A. R. (2025). "Ground-source heat pump performance in mixed-use urban developments." Applied Energy, 352, 121998.')
    doc.add_paragraph('European Commission. (2024). "Clean Energy for All Europeans: Urban Implementation Progress Report." Brussels: EC Directorate-General for Energy.')
    doc.add_paragraph('Fischer, H., & Nakamura, Y. (2024). "Wind energy potential in complex urban terrain." Journal of Wind Engineering, 248, 105091.')
    doc.add_paragraph('Garcia, R. M., & Thompson, D. W. (2025). "Property value impacts of distributed energy resources." Journal of Real Estate Economics, 52(1), 78-103.')
    doc.add_paragraph('International Energy Agency. (2025). "World Energy Outlook 2025: Cities and Energy Special Report." Paris: IEA Publications.')
    doc.add_paragraph('Johnson, B. A., & Lee, S. H. (2024). "Workforce development for the clean energy transition." Energy Research & Social Science, 98, 103012.')
    doc.add_paragraph('Kim, J., Park, C., & Singh, R. (2024). "District geothermal networks: Technical and economic assessment." Geothermics, 119, 102941.')
    doc.add_paragraph('Lawrence Berkeley National Laboratory. (2025). "Tracking the Sun: Distributed Solar Market Trends Report." Berkeley: LBNL.')

    doc.save(OUTPUT)
    print(f'Main document created: {OUTPUT}')


def create_template():
    """Create Academic_Report.ott template with custom styles via LibreOffice conversion.

    We first create a .docx with custom styles, then convert to .ott on the VM.
    """
    doc = Document()

    # Define custom Heading 1 style
    # Modify existing Heading 1
    h1_style = doc.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Georgia'
    h1_font.size = Pt(24)
    h1_font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)  # Dark navy blue
    h1_font.bold = True
    h1_pf = h1_style.paragraph_format
    h1_pf.space_before = Pt(24)
    h1_pf.space_after = Pt(12)

    # Modify existing Heading 2
    h2_style = doc.styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Georgia'
    h2_font.size = Pt(18)
    h2_font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)  # Medium blue
    h2_font.bold = True
    h2_font.italic = True
    h2_pf = h2_style.paragraph_format
    h2_pf.space_before = Pt(18)
    h2_pf.space_after = Pt(8)

    # Modify Body Text style
    # Need to add it if not present
    try:
        bt_style = doc.styles['Body Text']
    except KeyError:
        bt_style = doc.styles.add_style('Body Text', 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    bt_font = bt_style.font
    bt_font.name = 'Palatino Linotype'
    bt_font.size = Pt(11)
    bt_font.color.rgb = RGBColor(0x33, 0x33, 0x33)  # Dark gray
    bt_pf = bt_style.paragraph_format
    bt_pf.space_after = Pt(8)
    bt_pf.line_spacing = 1.15

    # Create/Modify Caption style
    try:
        cap_style = doc.styles['Caption']
    except KeyError:
        cap_style = doc.styles.add_style('Caption', 1)
    cap_font = cap_style.font
    cap_font.name = 'Georgia'
    cap_font.size = Pt(9)
    cap_font.italic = True
    cap_font.color.rgb = RGBColor(0x66, 0x66, 0x66)  # Medium gray
    cap_pf = cap_style.paragraph_format
    cap_pf.space_before = Pt(4)
    cap_pf.space_after = Pt(12)
    cap_pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Create a custom style unique to the template: "Academic Note"
    try:
        note_style = doc.styles['Academic Note']
    except KeyError:
        note_style = doc.styles.add_style('Academic Note', 1)
    note_font = note_style.font
    note_font.name = 'Georgia'
    note_font.size = Pt(10)
    note_font.italic = True
    note_font.color.rgb = RGBColor(0x8B, 0x45, 0x13)  # Saddle brown
    note_pf = note_style.paragraph_format
    note_pf.left_indent = Inches(0.5)
    note_pf.right_indent = Inches(0.5)
    note_pf.space_before = Pt(6)
    note_pf.space_after = Pt(6)

    # Create another custom style: "Academic Subtitle"
    try:
        sub_style = doc.styles['Academic Subtitle']
    except KeyError:
        sub_style = doc.styles.add_style('Academic Subtitle', 1)
    sub_font = sub_style.font
    sub_font.name = 'Georgia'
    sub_font.size = Pt(14)
    sub_font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    sub_font.italic = True
    sub_pf = sub_style.paragraph_format
    sub_pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_pf.space_before = Pt(6)
    sub_pf.space_after = Pt(18)

    # Add sample content to demonstrate styles
    doc.add_heading('Academic Report Template', level=1)
    doc.add_heading('Subtitle Section', level=2)
    p = doc.add_paragraph('This is body text in the Academic Report style.', style='Body Text')
    p = doc.add_paragraph('Caption: Sample figure caption text.', style='Caption')

    doc.save(TEMPLATE_DOCX)
    print(f'Template DOCX created: {TEMPLATE_DOCX}')


def convert_docx_to_ott():
    """Convert the template .docx to .ott using LibreOffice command line."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    env["HOME"] = "/home/user"

    # Convert .docx to .ott
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'ott',
         '--outdir', WORKDIR, TEMPLATE_DOCX],
        env=env,
        capture_output=True,
        text=True,
        timeout=30
    )
    print(f'Conversion stdout: {result.stdout}')
    print(f'Conversion stderr: {result.stderr}')

    # Rename from Academic_Report_temp.ott to Academic_Report.ott
    temp_ott = os.path.join(WORKDIR, 'Academic_Report_temp.ott')
    if os.path.exists(temp_ott):
        os.rename(temp_ott, TEMPLATE_OTT)
        print(f'Template OTT created: {TEMPLATE_OTT}')
    else:
        print(f'WARNING: Expected {temp_ott} not found after conversion')
        # List what was created
        for f in os.listdir(WORKDIR):
            if 'Academic' in f or f.endswith('.ott'):
                print(f'  Found: {WORKDIR}/{f}')

    # Clean up temp docx
    if os.path.exists(TEMPLATE_DOCX):
        os.remove(TEMPLATE_DOCX)
        print('Cleaned up temp DOCX')


# Main execution
create_main_document()
create_template()
convert_docx_to_ott()

# Launch LibreOffice Writer with the main document
launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')
