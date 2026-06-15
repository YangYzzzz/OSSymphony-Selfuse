"""
Initial Setup: Set up a Writer document with a custom page background/watermark.
Task ID: writer_rd_090
Domain: libreoffice_writer

Creates:
  - /home/user/watermark_logo.png (light gray company logo)
  - /home/user/writer_rd_090.docx (10-page corporate report, plain white background)
  - Opens the document in LibreOffice Writer
"""

import os
import shlex
import subprocess
import time

# Install dependencies on VM
subprocess.run(['pip3', 'install', 'python-docx', 'Pillow'], capture_output=True)

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_090'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
WATERMARK_IMG = f'{WORKDIR}/watermark_logo.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_watermark_logo():
    """Create a light gray company logo suitable for use as a watermark background."""
    from PIL import Image, ImageDraw, ImageFont

    width, height = 600, 600
    img = Image.new('RGBA', (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    # Draw a stylized company logo in light gray
    # Outer circle
    circle_color = (200, 200, 200, 80)
    draw.ellipse([50, 50, 550, 550], outline=circle_color, width=8)
    # Inner hexagon shape
    import math
    cx, cy, r = 300, 280, 150
    hex_points = [(cx + r * math.cos(math.radians(a)), cy + r * math.sin(math.radians(a)))
                  for a in range(0, 360, 60)]
    draw.polygon(hex_points, outline=circle_color, fill=None)
    for i in range(len(hex_points)):
        draw.line([hex_points[i], (cx, cy)], fill=circle_color, width=3)

    # Company name text
    text_color = (180, 180, 180, 100)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 48)
        font_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 24)
    except (OSError, IOError):
        font = ImageFont.load_default()
        font_small = font

    draw.text((300, 470), "NEXTERA", fill=text_color, font=font, anchor="mm")
    draw.text((300, 520), "CONSULTING GROUP", fill=text_color, font=font_small, anchor="mm")

    img.save(WATERMARK_IMG)
    print(f'Watermark logo created: {WATERMARK_IMG}')


def create_report():
    """Create a 10-page corporate report with plain white background."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    def add_heading_styled(text, level=1):
        h = doc.add_heading(text, level=level)
        return h

    def add_body_text(text):
        p = doc.add_paragraph(text)
        pf = p.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15
        return p

    # --- Page 1: Title Page ---
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_heading('NextEra Consulting Group', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Annual Strategic Review 2025')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = info.add_run('Prepared by the Office of the Chief Strategy Officer\nMarch 2025 | Confidential')
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # --- Page 2: Table of Contents ---
    add_heading_styled('Table of Contents', level=1)
    toc_items = [
        ('1. Executive Summary', '3'),
        ('2. Market Analysis', '4'),
        ('3. Financial Performance', '5'),
        ('4. Client Portfolio Overview', '6'),
        ('5. Operational Highlights', '7'),
        ('6. Technology & Innovation', '8'),
        ('7. Human Capital Report', '9'),
        ('8. Strategic Outlook 2026', '10'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{item}')
        run.font.size = Pt(12)
        dots = '.' * (60 - len(item))
        run2 = p.add_run(f' {dots} {page}')
        run2.font.size = Pt(12)
        run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # --- Page 3: Executive Summary ---
    add_heading_styled('1. Executive Summary', level=1)
    add_body_text(
        'NextEra Consulting Group delivered exceptional results in fiscal year 2025, '
        'achieving revenue growth of 18.7% year-over-year to reach $247.3 million in total revenue. '
        'This performance exceeded our initial projections by 4.2 percentage points, driven primarily '
        'by strong demand in digital transformation and cloud migration services.'
    )
    add_body_text(
        'Our client retention rate remained industry-leading at 94.2%, while new client acquisitions '
        'increased by 23% compared to the prior year. The firm successfully expanded into three new '
        'geographic markets: Singapore, Dubai, and Sao Paulo, establishing dedicated teams in each location.'
    )
    add_body_text(
        'Key strategic initiatives completed during the year include the acquisition of DataStream Analytics, '
        'a boutique AI consulting firm, the launch of our proprietary NextEra Insights Platform, and the '
        'establishment of the NextEra Innovation Lab in partnership with MIT Sloan School of Management.'
    )
    add_body_text(
        'Looking ahead to 2026, the leadership team has identified artificial intelligence advisory, '
        'sustainability consulting, and cybersecurity assessment as the three highest-priority growth vectors. '
        'We project consolidated revenue of $285-295 million for the coming fiscal year.'
    )
    doc.add_page_break()

    # --- Page 4: Market Analysis ---
    add_heading_styled('2. Market Analysis', level=1)
    add_heading_styled('2.1 Industry Landscape', level=2)
    add_body_text(
        'The global management consulting market reached $362 billion in 2025, representing a compound '
        'annual growth rate of 7.3% over the past five years. Digital transformation consulting continues '
        'to be the fastest-growing segment, accounting for approximately 28% of total market revenue.'
    )
    add_heading_styled('2.2 Competitive Positioning', level=2)
    add_body_text(
        'NextEra maintains a strong position in the mid-market enterprise segment, serving organizations '
        'with annual revenues between $500 million and $5 billion. Our Net Promoter Score of 72 places us '
        'in the top quartile among consulting firms of comparable size.'
    )
    add_body_text(
        'Primary competitive advantages include our integrated technology-strategy methodology, deep '
        'industry specialization in financial services and healthcare, and our talent retention rate of '
        '91.8%, which is 15 points above the industry average.'
    )
    add_heading_styled('2.3 Emerging Trends', level=2)
    add_body_text(
        'Three macro trends are reshaping client demand: (1) the accelerating adoption of generative AI '
        'across enterprise operations, (2) increasing regulatory complexity around ESG reporting and data '
        'privacy, and (3) the shift toward outcome-based pricing models in professional services engagements.'
    )
    doc.add_page_break()

    # --- Page 5: Financial Performance ---
    add_heading_styled('3. Financial Performance', level=1)
    add_body_text(
        'The following table summarizes key financial metrics for the past three fiscal years:'
    )

    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    headers = ['Metric', 'FY2023', 'FY2024', 'FY2025']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    fin_data = [
        ['Total Revenue ($M)', '$192.4', '$208.3', '$247.3'],
        ['Operating Income ($M)', '$38.5', '$43.7', '$54.4'],
        ['Operating Margin', '20.0%', '21.0%', '22.0%'],
        ['EBITDA ($M)', '$44.2', '$50.1', '$61.8'],
        ['Revenue per Consultant ($K)', '$312', '$328', '$357'],
        ['Utilization Rate', '74.2%', '76.1%', '78.5%'],
        ['New Engagements', '142', '168', '207'],
    ]
    for r, row_data in enumerate(fin_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    add_body_text('')
    add_body_text(
        'Revenue growth was primarily driven by the Digital Transformation practice (+31.2%), '
        'followed by Data & Analytics (+22.8%) and Strategy Advisory (+12.4%). The firm achieved '
        'positive operating leverage, with operating margin expanding 100 basis points to 22.0%.'
    )
    doc.add_page_break()

    # --- Page 6: Client Portfolio ---
    add_heading_styled('4. Client Portfolio Overview', level=1)
    add_body_text(
        'NextEra served 189 active clients across 14 industry verticals during FY2025. '
        'Our top 20 clients accounted for 43% of total revenue, compared to 48% in the prior year, '
        'reflecting healthy diversification of our client base.'
    )

    table2 = doc.add_table(rows=7, cols=3)
    table2.style = 'Table Grid'
    for i, h in enumerate(['Industry Vertical', 'Revenue Share', '# of Clients']):
        cell = table2.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    client_data = [
        ['Financial Services', '28.4%', '42'],
        ['Healthcare & Life Sciences', '19.7%', '31'],
        ['Technology & Media', '16.3%', '28'],
        ['Manufacturing & Industrial', '14.1%', '35'],
        ['Energy & Utilities', '11.8%', '27'],
        ['Retail & Consumer Goods', '9.7%', '26'],
    ]
    for r, row_data in enumerate(client_data, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    add_body_text('')
    add_body_text(
        'Notable new client wins in 2025 include Meridian Financial Holdings ($4.2M engagement), '
        'Pacific Northwest Health System ($3.8M), and Vertex Semiconductor ($2.9M). '
        'These wins validate our sector-focused go-to-market strategy.'
    )
    doc.add_page_break()

    # --- Page 7: Operational Highlights ---
    add_heading_styled('5. Operational Highlights', level=1)
    add_body_text(
        'The firm invested significantly in operational infrastructure during 2025 to support '
        'continued scaling. Major initiatives included:'
    )
    items = [
        'Implementation of the Salesforce Revenue Cloud platform for engagement management and billing',
        'Migration of all internal systems to AWS with 99.97% uptime achieved',
        'Launch of the NextEra Knowledge Hub, an internal AI-powered research assistant',
        'Rollout of standardized project governance framework across all practice areas',
        'Establishment of a centralized Proposal Support Center reducing response time by 40%',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    add_body_text('')
    add_body_text(
        'These operational investments are projected to yield $8.5 million in annual efficiency '
        'gains beginning in FY2026, primarily through reduced administrative overhead and improved '
        'consultant utilization rates.'
    )
    doc.add_page_break()

    # --- Page 8: Technology & Innovation ---
    add_heading_styled('6. Technology & Innovation', level=1)
    add_heading_styled('6.1 NextEra Insights Platform', level=2)
    add_body_text(
        'The NextEra Insights Platform, launched in Q2 2025, is our proprietary analytics tool that '
        'enables clients to benchmark their digital maturity against industry peers. The platform '
        'currently serves 47 enterprise clients and has generated $12.3 million in recurring SaaS revenue.'
    )
    add_heading_styled('6.2 AI Advisory Practice', level=2)
    add_body_text(
        'Our dedicated AI Advisory practice grew from 12 to 38 consultants during the year, reflecting '
        'surging demand for responsible AI governance, large language model strategy, and AI-powered '
        'process automation. The practice generated $18.7 million in revenue, a 156% increase YoY.'
    )
    add_heading_styled('6.3 Innovation Lab', level=2)
    add_body_text(
        'The NextEra Innovation Lab, established in partnership with MIT Sloan, has produced three '
        'proprietary frameworks for AI readiness assessment, sustainability impact measurement, and '
        'workforce transformation planning. These frameworks are being integrated into client engagements.'
    )
    doc.add_page_break()

    # --- Page 9: Human Capital ---
    add_heading_styled('7. Human Capital Report', level=1)
    add_body_text(
        'NextEra employs 692 professionals across 8 offices globally. Key HR metrics for 2025:'
    )

    table3 = doc.add_table(rows=7, cols=3)
    table3.style = 'Table Grid'
    for i, h in enumerate(['Metric', 'FY2024', 'FY2025']):
        cell = table3.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    hr_data = [
        ['Total Headcount', '583', '692'],
        ['Voluntary Turnover', '12.4%', '8.2%'],
        ['Gender Diversity (Women %)', '38%', '42%'],
        ['Avg Training Hours/Employee', '64', '78'],
        ['Employee Engagement Score', '4.1/5.0', '4.4/5.0'],
        ['Campus Hires', '45', '62'],
    ]
    for r, row_data in enumerate(hr_data, 1):
        for c, val in enumerate(row_data):
            table3.cell(r, c).text = val

    add_body_text('')
    add_body_text(
        'The firm launched the NextEra Leadership Academy in 2025, a 12-month development program '
        'for high-potential managers. The inaugural cohort of 24 participants completed the program '
        'in December, with 8 participants receiving promotions to Principal or Partner level.'
    )
    doc.add_page_break()

    # --- Page 10: Strategic Outlook ---
    add_heading_styled('8. Strategic Outlook 2026', level=1)
    add_body_text(
        'The leadership team has established three strategic pillars for fiscal year 2026:'
    )
    add_heading_styled('Pillar 1: Intelligent Enterprise', level=2)
    add_body_text(
        'Expand AI advisory capabilities to help clients implement responsible AI at scale. '
        'Target: grow AI practice revenue to $35 million.'
    )
    add_heading_styled('Pillar 2: Sustainable Growth', level=2)
    add_body_text(
        'Launch dedicated ESG and sustainability consulting practice. Recruit 15 senior specialists '
        'with deep regulatory expertise in climate disclosure and carbon accounting.'
    )
    add_heading_styled('Pillar 3: Global Reach', level=2)
    add_body_text(
        'Establish presence in two additional markets (Tokyo and Frankfurt) to serve existing '
        'multinational clients and capture local demand. Target: 15% of revenue from international offices.'
    )
    add_body_text('')
    add_body_text(
        'Consolidated revenue guidance for FY2026 is $285-295 million, representing 15-19% growth. '
        'Operating margin is expected to remain stable at 21.5-22.5% as incremental investment in '
        'new practices offsets operational efficiency gains.'
    )
    add_body_text(
        'The Board of Directors has approved a $15 million investment budget for strategic acquisitions '
        'and technology platform development, with a focus on bolt-on acquisitions that enhance our '
        'capabilities in cybersecurity consulting and data engineering.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')


def main():
    create_watermark_logo()
    create_report()
    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
