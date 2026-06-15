"""
Initial Setup: Create a 20-page thesis document with no page numbering
Task ID: writer_tm_071
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_071'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def fill_page_with_text(doc, paragraphs_list, space_after=Pt(6)):
    """Add multiple paragraphs to fill space on a page."""
    for text in paragraphs_list:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = space_after
        for run in para.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'


def create_initial():
    doc = Document()

    # Set default page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ===== PAGE 1: Title Page =====
    # Add some blank space before title
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('Advancing Neural Network Architectures\nfor Climate Pattern Recognition')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_before = Pt(36)
    run = subtitle.add_run('A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    author = doc.add_paragraph()
    author.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_before = Pt(48)
    run = author.add_run('By\nDr. Elena Vasquez Rodriguez')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    dept = doc.add_paragraph()
    dept.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept.paragraph_format.space_before = Pt(24)
    run = dept.add_run('Department of Computer Science\nStanford University\nMarch 2025')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # ===== PAGE 2: Abstract =====
    doc.add_page_break()

    heading = doc.add_paragraph()
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading.paragraph_format.space_after = Pt(18)
    run = heading.add_run('ABSTRACT')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    abstract_text = [
        'This dissertation presents a comprehensive investigation into the application of advanced neural network architectures for recognizing and predicting complex climate patterns. Our research addresses the critical challenge of processing multi-dimensional atmospheric data to identify emerging trends in global temperature distributions, precipitation cycles, and extreme weather events.',
        'We introduce ClimateNet, a novel hybrid architecture combining transformer attention mechanisms with graph neural networks, specifically designed for spatiotemporal climate data analysis. The model processes satellite imagery, ground-based sensor readings, and historical meteorological records to generate accurate 30-day forecasts with unprecedented spatial resolution.',
        'Our experimental evaluation demonstrates that ClimateNet achieves a 23.7% improvement in prediction accuracy compared to existing state-of-the-art methods across five benchmark datasets. The model exhibits particular strength in identifying El Nino Southern Oscillation precursors and Arctic vortex disruption patterns, achieving recall rates of 94.2% and 89.8% respectively.',
        'Furthermore, we present a novel attention-based interpretability framework that enables climate scientists to understand which input features most strongly influence model predictions. This transparency mechanism has been validated through collaboration with domain experts at the National Oceanic and Atmospheric Administration (NOAA).',
        'The contributions of this work span three key areas: (1) a scalable architecture for climate pattern recognition, (2) a comprehensive benchmark suite for evaluating climate prediction models, and (3) an interpretability framework bridging machine learning and atmospheric science.',
    ]
    fill_page_with_text(doc, abstract_text)

    # ===== PAGE 3: Table of Contents =====
    doc.add_page_break()

    heading = doc.add_paragraph()
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading.paragraph_format.space_after = Pt(24)
    run = heading.add_run('TABLE OF CONTENTS')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    toc_entries = [
        ('Chapter 1: Introduction', '1'),
        ('   1.1 Background and Motivation', '1'),
        ('   1.2 Research Questions', '3'),
        ('   1.3 Contributions', '4'),
        ('   1.4 Thesis Organization', '5'),
        ('Chapter 2: Literature Review', '6'),
        ('   2.1 Climate Modeling Approaches', '6'),
        ('   2.2 Deep Learning for Geospatial Data', '8'),
        ('   2.3 Transformer Architectures', '10'),
        ('   2.4 Graph Neural Networks', '11'),
        ('Chapter 3: Methodology', '13'),
        ('   3.1 ClimateNet Architecture', '13'),
        ('   3.2 Data Preprocessing Pipeline', '15'),
        ('   3.3 Training Procedure', '16'),
        ('Chapter 4: Experimental Evaluation', '18'),
        ('   4.1 Datasets and Benchmarks', '18'),
        ('   4.2 Baseline Comparisons', '20'),
        ('   4.3 Ablation Studies', '22'),
        ('Chapter 5: Interpretability Framework', '24'),
        ('Chapter 6: Discussion and Future Work', '26'),
        ('Chapter 7: Conclusion', '28'),
        ('References', '30'),
    ]

    for entry, page in toc_entries:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(f'{entry}')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run2 = para.add_run(f'  {"." * (50 - len(entry))}  {page}')
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'

    # ===== PAGES 4-20: Main Content (17 pages) =====
    # We need to generate enough content for 17 pages

    chapters = [
        {
            'title': 'Chapter 1: Introduction',
            'sections': [
                ('1.1 Background and Motivation',
                 [
                     'Climate change represents one of the most pressing challenges facing humanity in the twenty-first century. The ability to accurately predict and understand complex atmospheric patterns is crucial for developing effective mitigation and adaptation strategies. Traditional numerical weather prediction models, while valuable, are computationally expensive and often struggle to capture the nonlinear dynamics inherent in climate systems.',
                     'Recent advances in deep learning have demonstrated remarkable success in pattern recognition across various domains, from computer vision to natural language processing. These developments have inspired researchers to explore neural network-based approaches for climate modeling, potentially complementing and enhancing existing physical models.',
                     'The intersection of artificial intelligence and atmospheric science presents unique opportunities and challenges. Unlike standard image classification or text generation tasks, climate data is inherently spatiotemporal, multi-scale, and governed by well-understood physical laws. Any successful machine learning approach must respect these constraints while leveraging the flexibility of data-driven methods.',
                 ]),
                ('1.2 Research Questions',
                 [
                     'This dissertation addresses three primary research questions that guide our investigation:',
                     'First, can a hybrid neural architecture combining transformer attention mechanisms with graph neural networks effectively capture both local and global dependencies in spatiotemporal climate data? We hypothesize that attention mechanisms can identify long-range teleconnections while graph networks encode the physical topology of observation stations.',
                     'Second, does explicit incorporation of physical constraints through custom loss functions improve prediction accuracy compared to purely data-driven approaches? Previous work has shown that physics-informed neural networks can achieve better generalization, but their application to climate prediction remains underexplored.',
                     'Third, can we develop interpretability methods that provide meaningful explanations of model predictions to domain scientists? The black-box nature of deep learning models has been a significant barrier to their adoption in operational climate forecasting.',
                 ]),
                ('1.3 Contributions',
                 [
                     'The primary contributions of this dissertation are threefold. We present ClimateNet, a novel hybrid architecture specifically designed for spatiotemporal climate data analysis. The architecture combines multi-head self-attention layers with graph convolution modules, enabling the model to process irregularly sampled observation data while maintaining awareness of spatial relationships.',
                     'We introduce the Climate Prediction Benchmark Suite (CPBS), a standardized evaluation framework comprising five diverse datasets spanning different geographic regions, temporal scales, and climate phenomena. This benchmark addresses the current lack of standardized evaluation protocols in climate machine learning research.',
                     'Finally, we develop an attention-based interpretability framework that generates human-readable explanations of model predictions. Through collaboration with climate scientists at NOAA, we validate that these explanations align with established physical understanding of atmospheric dynamics.',
                 ]),
            ],
        },
        {
            'title': 'Chapter 2: Literature Review',
            'sections': [
                ('2.1 Climate Modeling Approaches',
                 [
                     'The history of computational climate modeling spans over six decades, beginning with the pioneering work of Phillips (1956) on general circulation models. Modern climate models, such as the Community Earth System Model (CESM) and the Geophysical Fluid Dynamics Laboratory (GFDL) model, simulate the coupled ocean-atmosphere system using numerical solutions to partial differential equations governing fluid dynamics, thermodynamics, and radiative transfer.',
                     'Despite their physical fidelity, these models face significant computational limitations. A single century-long simulation at high spatial resolution can require millions of CPU hours on modern supercomputers. This computational burden restricts the number of ensemble members that can be generated, limiting uncertainty quantification in climate projections.',
                     'Statistical downscaling methods have emerged as a complement to dynamical models, using empirical relationships between large-scale atmospheric variables and local climate conditions. While computationally efficient, these methods assume stationarity of statistical relationships, an assumption increasingly challenged by the non-stationary nature of climate change.',
                 ]),
                ('2.2 Deep Learning for Geospatial Data',
                 [
                     'The application of deep learning to geospatial data has grown rapidly since the seminal work of Krizhevsky et al. (2012) on convolutional neural networks. Early applications focused on satellite image classification, land use mapping, and object detection in remote sensing imagery. These approaches typically treated geographic data as standard images, applying two-dimensional convolutions to gridded representations.',
                     'More recent work has recognized the unique properties of geospatial data that distinguish it from natural images. Geographic data often exhibits strong spatial autocorrelation, non-Euclidean geometry on the sphere, and multi-resolution characteristics that require specialized architectural considerations. Graph neural networks and spherical convolutions have been proposed to address these challenges.',
                     'In the context of weather and climate prediction, convolutional architectures have been applied to gridded reanalysis data with promising results. Weyn et al. (2020) demonstrated that a CNN-based model could produce skillful medium-range weather forecasts, while Rasp et al. (2020) showed that neural networks could effectively parameterize subgrid-scale processes in climate models.',
                 ]),
                ('2.3 Transformer Architectures',
                 [
                     'The transformer architecture, introduced by Vaswani et al. (2017) for machine translation, has revolutionized multiple fields of deep learning. The self-attention mechanism at its core enables direct modeling of long-range dependencies without the sequential processing constraints of recurrent neural networks.',
                     'Vision transformers (ViT) adapted the architecture for image processing by treating images as sequences of patches. This approach has been extended to video understanding, point cloud processing, and various scientific applications. The ability to capture global context from the first layer makes transformers particularly attractive for climate data, where teleconnections spanning thousands of kilometers play crucial roles.',
                     'Recent applications of transformers to weather prediction have achieved remarkable success. Pathak et al. (2022) developed FourCastNet, a Fourier neural operator-based model achieving competitive performance with operational numerical weather prediction systems at a fraction of the computational cost.',
                 ]),
            ],
        },
        {
            'title': 'Chapter 3: Methodology',
            'sections': [
                ('3.1 ClimateNet Architecture',
                 [
                     'ClimateNet consists of three main components: an encoder module that processes multi-modal input data, a spatiotemporal processing core that captures complex dependencies, and a decoder module that generates predictions at the desired spatial and temporal resolution.',
                     'The encoder module handles three types of input: gridded satellite imagery at 0.25-degree resolution, point-based station observations from 12,847 meteorological stations worldwide, and auxiliary categorical features including terrain type, land cover classification, and proximity to water bodies. Each modality is processed through a dedicated embedding layer before fusion.',
                     'The spatiotemporal core alternates between graph attention layers and temporal transformer layers. Graph attention operates on a dynamically constructed graph where nodes represent spatial locations and edges encode proximity and physical relationships. Temporal transformer layers apply multi-head self-attention across the time dimension, enabling the model to identify relevant historical patterns across varying lag periods.',
                     'We incorporate physical constraints through a custom loss function that penalizes predictions violating conservation of energy and mass. This physics-informed regularization term is weighted by a hyperparameter lambda, tuned through cross-validation on the development set.',
                 ]),
                ('3.2 Data Preprocessing Pipeline',
                 [
                     'Our preprocessing pipeline addresses the heterogeneous nature of climate data through several stages. Raw satellite data from ERA5 reanalysis is regridded to a uniform 0.25-degree resolution using bilinear interpolation. Missing values, which occur primarily over polar regions and at satellite swath boundaries, are imputed using a combination of spatial interpolation and temporal persistence.',
                     'Station observations undergo quality control following World Meteorological Organization guidelines. We apply range checks, temporal consistency tests, and spatial buddy checks to identify and flag erroneous measurements. Flagged values are replaced using inverse distance weighted interpolation from neighboring stations within a 200-kilometer radius.',
                     'Feature normalization is performed independently for each variable using the mean and standard deviation computed from the training period (1979-2015). Temperature fields are additionally detrended using a linear fit to the global mean temperature time series, following the approach of Watson-Parris et al. (2021).',
                 ]),
                ('3.3 Training Procedure',
                 [
                     'ClimateNet is trained using a two-phase approach. In the first phase, the model is trained on one-step prediction with a combined loss function comprising mean squared error for continuous variables and cross-entropy for categorical outputs. The physics-informed regularization term is gradually increased from zero to its full weight over the first 50 epochs using a linear warm-up schedule.',
                     'The second phase employs curriculum learning, progressively increasing the forecast lead time from 6 hours to 30 days. At each curriculum stage, the model is fine-tuned for 20 epochs using the AdamW optimizer with a cosine learning rate schedule. This approach helps the model learn to accumulate errors gracefully and produce physically consistent long-range forecasts.',
                     'Training is distributed across 32 NVIDIA A100 GPUs using data parallelism with gradient accumulation. The total training time for the full curriculum is approximately 72 hours. We use mixed-precision training to reduce memory consumption and enable larger batch sizes.',
                 ]),
            ],
        },
        {
            'title': 'Chapter 4: Experimental Evaluation',
            'sections': [
                ('4.1 Datasets and Benchmarks',
                 [
                     'We evaluate ClimateNet on the Climate Prediction Benchmark Suite (CPBS), comprising five datasets designed to test different aspects of climate prediction capability. The Global Temperature Prediction dataset covers monthly mean temperature anomalies at 2.5-degree resolution from 1950 to 2023, targeting predictions at 1, 3, 6, and 12-month lead times.',
                     'The Regional Precipitation dataset focuses on daily precipitation over the continental United States at 0.25-degree resolution, sourced from the PRISM dataset. This challenging benchmark tests the model ability to predict highly variable and spatially heterogeneous precipitation patterns.',
                     'The Extreme Weather Events dataset comprises labeled instances of tropical cyclones, atmospheric rivers, and heat waves from 1990 to 2022. This dataset evaluates the model ability to detect and track extreme events, which are of particular societal importance.',
                     'The ENSO Prediction dataset targets the Nino 3.4 index at lead times from 1 to 24 months, following the standard evaluation protocol established by Ham et al. (2019). The Arctic Oscillation dataset evaluates prediction of the Arctic Oscillation index, an important mode of atmospheric variability affecting mid-latitude weather patterns.',
                 ]),
                ('4.2 Baseline Comparisons',
                 [
                     'We compare ClimateNet against several baselines spanning traditional statistical methods, dynamical models, and state-of-the-art machine learning approaches. Statistical baselines include persistence, climatology, and vector autoregression models. Dynamical model predictions are obtained from the ECMWF Integrated Forecasting System (IFS) and the NCEP Climate Forecast System version 2 (CFSv2).',
                     'Machine learning baselines include ConvLSTM (Shi et al., 2015), FourCastNet (Pathak et al., 2022), Pangu-Weather (Bi et al., 2023), and GraphCast (Lam et al., 2023). All models are trained on the same data splits using their published hyperparameter configurations.',
                     'Results demonstrate that ClimateNet achieves state-of-the-art performance on four of the five benchmark datasets. On the Global Temperature Prediction dataset, ClimateNet achieves a root mean squared error (RMSE) of 0.287 degrees Celsius at the 6-month lead time, compared to 0.376 for GraphCast and 0.412 for ConvLSTM. The improvement is particularly pronounced at longer lead times.',
                 ]),
                ('4.3 Ablation Studies',
                 [
                     'We conduct systematic ablation studies to quantify the contribution of each architectural component. Removing the graph attention layers while retaining the transformer backbone reduces performance by 8.3% on the Global Temperature dataset, confirming the importance of explicitly modeling spatial relationships.',
                     'Ablating the physics-informed loss term leads to a 5.1% degradation in RMSE but a more substantial 12.7% increase in physical inconsistency metrics, measured by conservation law violations in the predictions. This result supports our hypothesis that physical constraints improve not only accuracy but also the physical plausibility of predictions.',
                     'The multi-modal encoder contributes differently depending on the target variable. For temperature prediction, station observations provide the largest marginal improvement (4.2%), while for precipitation, satellite imagery contributes most significantly (7.8%). This finding suggests that the optimal input configuration may vary by application.',
                 ]),
            ],
        },
        {
            'title': 'Chapter 5: Interpretability Framework',
            'sections': [
                ('5.1 Attention-Based Explanations',
                 [
                     'Our interpretability framework leverages the attention weights computed by the transformer layers to generate spatial and temporal attribution maps. For each prediction, we extract attention patterns from all heads and layers, then aggregate them using a gradient-weighted approach that accounts for the relative importance of different attention heads.',
                     'Spatial attribution maps highlight geographic regions that most strongly influence the prediction for a given target location. Temporal attribution maps identify which historical time steps carry the most predictive signal. Together, these visualizations provide climate scientists with intuitive explanations of model behavior.',
                     'We validate the interpretability framework through a user study involving 15 climate scientists from NOAA and three university research groups. Participants rated the relevance and usefulness of model explanations on a 5-point Likert scale. The mean relevance score was 4.2, with scientists finding spatial attributions particularly useful for identifying teleconnection patterns consistent with known physical mechanisms.',
                 ]),
                ('5.2 Case Studies',
                 [
                     'We present three detailed case studies demonstrating the utility of our interpretability framework. The first examines ClimateNet predictions for the 2015-2016 El Nino event, one of the strongest on record. Attention maps reveal that the model correctly identifies the buildup of warm subsurface waters in the equatorial Pacific as a key precursor signal, consistent with established ENSO theory.',
                     'The second case study analyzes the February 2021 Arctic cold outbreak that caused widespread disruption across the southern United States. Our model attention maps highlight the weakening of the polar vortex two weeks prior and the subsequent southward displacement of the jet stream, aligning with the meteorological analysis published by Cohen et al. (2021).',
                     'The third case study explores the model predictions for the unprecedented Pacific Northwest heat dome of June 2021. The interpretability framework identifies anomalous blocking patterns in the mid-troposphere and unusually warm sea surface temperatures as primary driving factors, consistent with the post-event attribution study by Philip et al. (2022).',
                 ]),
            ],
        },
        {
            'title': 'Chapter 6: Discussion and Future Work',
            'sections': [
                ('6.1 Limitations',
                 [
                     'Despite the promising results presented in this dissertation, several limitations warrant discussion. First, ClimateNet performance degradation beyond the 30-day forecast horizon suggests that the current architecture struggles to capture the slower modes of variability that dominate at seasonal and longer time scales. Incorporating ocean dynamics more explicitly into the model may help address this limitation.',
                     'Second, the computational cost of training, while substantially less than running a full climate model, remains significant for individual researchers. The requirement for multi-GPU infrastructure limits accessibility and reproducibility. Developing efficient model compression and distillation techniques is an important direction for future work.',
                     'Third, our evaluation has focused primarily on midlatitude and tropical regions where observation networks are dense. Model performance in data-sparse regions, particularly over the Southern Ocean and polar areas, remains uncertain and requires dedicated evaluation campaigns.',
                 ]),
                ('6.2 Future Directions',
                 [
                     'Several promising research directions emerge from this work. The integration of ClimateNet with traditional numerical models in a hybrid framework could leverage the complementary strengths of both approaches. Early experiments with a coupled system, where ClimateNet provides subgrid-scale parameterizations within a coarse-resolution dynamical model, show encouraging results.',
                     'Extending the framework to support multi-decadal climate projections under different emission scenarios would significantly broaden its applicability. This extension requires careful consideration of distributional shift, as future climate conditions may differ substantially from the training period.',
                     'Finally, developing online learning capabilities that allow the model to continuously assimilate new observations could enable real-time updating of predictions as new data becomes available. This capability would be particularly valuable for operational forecasting applications.',
                 ]),
            ],
        },
        {
            'title': 'Chapter 7: Conclusion',
            'sections': [
                ('7.1 Summary of Contributions',
                 [
                     'This dissertation has presented ClimateNet, a novel hybrid neural network architecture for climate pattern recognition and prediction. Through extensive experimentation on the Climate Prediction Benchmark Suite, we have demonstrated that combining transformer attention mechanisms with graph neural networks yields significant improvements over existing approaches.',
                     'The physics-informed training procedure ensures that model predictions respect fundamental conservation laws while maintaining the flexibility to capture complex nonlinear dynamics. Our interpretability framework bridges the gap between machine learning and atmospheric science, providing climate scientists with actionable insights into model behavior.',
                     'The Climate Prediction Benchmark Suite itself represents a significant contribution to the field, providing a standardized evaluation framework that enables rigorous comparison of different approaches. We hope this benchmark will accelerate progress in climate machine learning research.',
                 ]),
                ('7.2 Broader Impact',
                 [
                     'Accurate climate prediction has far-reaching implications for agriculture, water resource management, disaster preparedness, and energy planning. By improving prediction accuracy and providing interpretable explanations, ClimateNet can help decision-makers better anticipate and prepare for climate-related risks.',
                     'We acknowledge the potential for misuse of climate prediction tools, including over-reliance on model outputs without proper uncertainty quantification. We advocate for responsible deployment practices that communicate prediction confidence alongside point forecasts and encourage ensemble-based approaches that capture the inherent uncertainty in the climate system.',
                 ]),
            ],
        },
    ]

    for ch_idx, chapter in enumerate(chapters):
        doc.add_page_break()

        # Chapter title
        ch_heading = doc.add_paragraph()
        ch_heading.paragraph_format.space_after = Pt(18)
        run = ch_heading.add_run(chapter['title'])
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Times New Roman'

        for sec_title, sec_paragraphs in chapter['sections']:
            # Section heading
            sec_heading = doc.add_paragraph()
            sec_heading.paragraph_format.space_before = Pt(12)
            sec_heading.paragraph_format.space_after = Pt(6)
            run = sec_heading.add_run(sec_title)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'

            for text in sec_paragraphs:
                para = doc.add_paragraph(text)
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.first_line_indent = Inches(0.5)
                for run in para.runs:
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
