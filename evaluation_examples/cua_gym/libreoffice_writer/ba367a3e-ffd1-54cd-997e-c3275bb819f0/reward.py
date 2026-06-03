"""
Reward Script: Insert page number fields in footer with different formats
Task ID: writer_tm_071
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Document has 2 sections (section break separating front matter from main content)
  Component 2 (0.25): Section 0 page number format is lowerRoman starting at 1
  Component 3 (0.25): Section 0 footer contains PAGE field with roman numeral formatting
  Component 4 (0.15): Section 1 page number format is decimal starting at 1
  Component 5 (0.15): Section 1 footer contains PAGE field for Arabic numbering
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_071'


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice edits via Ctrl+S."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)
    print(f"INFO: Document has {num_sections} section(s)")

    # Component 1: Document has 2 sections (0.20 points)
    # Initial has 1 section; golden has 2 (section break between front matter and main content)
    try:
        if num_sections >= 2:
            print(f"PASS: Component 1 — Document has {num_sections} sections (>= 2) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Expected >= 2 sections, found {num_sections}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Gate: need at least 2 sections for remaining checks
    if num_sections < 2:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Component 2: Section 0 has lowerRoman page number format with start (0.25 points)
    # Initial has no pgNumType; golden has fmt=lowerRoman, start=1
    try:
        sect0 = doc.sections[0]
        sectPr0 = sect0._sectPr
        pgNumType0 = sectPr0.find(qn('w:pgNumType'))
        if pgNumType0 is not None:
            fmt0 = pgNumType0.get(qn('w:fmt'))
            start0 = pgNumType0.get(qn('w:start'))
            if fmt0 == 'lowerRoman':
                print(f"PASS: Component 2 — Section 0 pgNumType fmt=lowerRoman, start={start0} (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — Section 0 pgNumType fmt={fmt0}, expected lowerRoman")
        else:
            print(f"FAIL: Component 2 — Section 0 has no pgNumType element")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Section 0 footer contains PAGE field with roman numeral formatting (0.25 points)
    # Initial footer is empty/linked; golden has PAGE field with \* roman or similar
    try:
        sect0_footer = doc.sections[0].footer
        footer0_has_page_field = False
        footer0_has_roman = False

        for p in sect0_footer.paragraphs:
            for elem in p._element.iter():
                tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                if tag == 'instrText':
                    instr = elem.text or ''
                    if 'PAGE' in instr.upper():
                        footer0_has_page_field = True
                    if 'roman' in instr.lower():
                        footer0_has_roman = True

        if footer0_has_page_field and footer0_has_roman:
            print(f"PASS: Component 3 — Section 0 footer has PAGE field with roman format (0.25 pts)")
            total_score += 0.25
        elif footer0_has_page_field:
            # Has PAGE field but no roman switch — partial credit
            print(f"PARTIAL: Component 3 — Section 0 footer has PAGE field but no roman format switch (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — Section 0 footer missing PAGE field (has_page={footer0_has_page_field}, has_roman={footer0_has_roman})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Section 1 has decimal page number format with start=1 (0.15 points)
    # Initial has no section 1; golden section 1 has fmt=decimal, start=1
    try:
        sect1 = doc.sections[1]
        sectPr1 = sect1._sectPr
        pgNumType1 = sectPr1.find(qn('w:pgNumType'))
        if pgNumType1 is not None:
            fmt1 = pgNumType1.get(qn('w:fmt'))
            start1 = pgNumType1.get(qn('w:start'))
            if fmt1 == 'decimal' and start1 == '1':
                print(f"PASS: Component 4 — Section 1 pgNumType fmt=decimal, start=1 (0.15 pts)")
                total_score += 0.15
            elif fmt1 == 'decimal':
                print(f"PARTIAL: Component 4 — Section 1 fmt=decimal but start={start1} (expected 1) (0.08 pts)")
                total_score += 0.08
            else:
                print(f"FAIL: Component 4 — Section 1 pgNumType fmt={fmt1}, expected decimal")
        else:
            print(f"FAIL: Component 4 — Section 1 has no pgNumType element")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Section 1 footer contains PAGE field for Arabic numbering (0.15 points)
    # Initial has no section 1 footer; golden has PAGE field
    try:
        sect1_footer = doc.sections[1].footer
        footer1_has_page_field = False

        for p in sect1_footer.paragraphs:
            for elem in p._element.iter():
                tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                if tag == 'instrText':
                    instr = elem.text or ''
                    if 'PAGE' in instr.upper():
                        footer1_has_page_field = True

        if footer1_has_page_field:
            print(f"PASS: Component 5 — Section 1 footer has PAGE field (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — Section 1 footer missing PAGE field")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
