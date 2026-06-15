"""
Initial Setup: Document with an image placed at an arbitrary (non-centered) position
Task ID: writer_obj_036
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
import io
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from PIL import Image, ImageDraw, ImageFont

WORKDIR = '/home/user'
TASK_ID = 'centered_image_doc'
OUTPUT = f'{WORKDIR}/Desktop/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_sample_image_bytes(width_px=320, height_px=200):
    """Create a simple PNG image in memory."""
    img = Image.new('RGB', (width_px, height_px), color=(100, 149, 237))  # cornflower blue
    draw = ImageDraw.Draw(img)
    # Draw a simple graphic
    draw.rectangle([20, 20, width_px - 20, height_px - 20], outline=(255, 255, 255), width=3)
    draw.ellipse([60, 40, width_px - 60, height_px - 40], fill=(255, 200, 50), outline=(200, 150, 0), width=2)
    draw.line([20, 20, width_px - 20, height_px - 20], fill=(255, 255, 255), width=2)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def add_floating_image_arbitrary(doc, img_bytes, width_cm=8.0, height_cm=5.0):
    """
    Add a floating image with an arbitrary position (left-aligned, 4cm from top).
    This uses wp:anchor XML element for floating placement.
    Horizontal: left-aligned (posOffset from left, NOT centered)
    Vertical: 4cm from top of page (NOT 2cm)
    """
    # EMU conversions
    EMU_PER_CM = 914400 / 100  # 9144 EMU per mm, 914400 per inch
    EMU_PER_CM = 360000  # 1 cm = 360000 EMU
    width_emu = int(width_cm * EMU_PER_CM)
    height_emu = int(height_cm * EMU_PER_CM)

    # Arbitrary position: 1.5 cm from left, 4.0 cm from top (NOT centered, NOT 2cm)
    pos_x_emu = int(1.5 * EMU_PER_CM)   # 1.5 cm from left edge of page
    pos_y_emu = int(4.0 * EMU_PER_CM)   # 4.0 cm from top of page

    # Add image to document relationships
    para = doc.add_paragraph()
    run = para.add_run()

    # Add picture inline first to get the relationship ID
    pic = run.add_picture(img_bytes, width=Cm(width_cm), height=Cm(height_cm))

    # Get the inline element and convert to anchor
    inline = run._element.find('.//' + qn('wp:inline'))
    if inline is None:
        return

    # Build the anchor XML element
    anchor_xml = f'''<wp:anchor
        xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        distT="0" distB="0" distL="114300" distR="114300"
        simplePos="0" relativeHeight="251658240" behindDoc="0"
        locked="0" layoutInCell="1" allowOverlap="1">
  <wp:simplePos x="0" y="0"/>
  <wp:positionH relativeFrom="page">
    <wp:posOffset>{pos_x_emu}</wp:posOffset>
  </wp:positionH>
  <wp:positionV relativeFrom="page">
    <wp:posOffset>{pos_y_emu}</wp:posOffset>
  </wp:positionV>
  <wp:extent cx="{width_emu}" cy="{height_emu}"/>
  <wp:effectExtent l="0" t="0" r="0" b="0"/>
  <wp:wrapSquare wrapText="bothSides"/>
  <wp:docPr id="1" name="Picture 1"/>
  <wp:cNvGraphicFramePr>
    <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>
  </wp:cNvGraphicFramePr>
</wp:anchor>'''

    anchor_elem = etree.fromstring(anchor_xml)

    # Copy the graphic data from the inline element into anchor
    graphic = inline.find('.//' + qn('a:graphic'))
    if graphic is not None:
        anchor_elem.append(graphic)

    # Replace inline with anchor in the run element
    inline.getparent().replace(inline, anchor_elem)


def create_initial():
    doc = Document()

    # Page setup: standard A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # --- Title ---
    title_para = doc.add_heading('Annual Product Review 2024', level=1)

    # --- Introduction paragraph ---
    intro = doc.add_paragraph(
        'This report summarizes the key findings from our annual product review conducted '
        'across all regional offices. The data presented here reflects performance metrics '
        'from Q1 through Q4 of the fiscal year 2024.'
    )

    # --- Add floating image with arbitrary (non-centered, not 2cm from top) position ---
    img_bytes = create_sample_image_bytes(320, 200)
    add_floating_image_arbitrary(doc, img_bytes, width_cm=8.0, height_cm=5.0)

    # --- More content paragraphs ---
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'Revenue growth exceeded expectations by 12.4% year-over-year, driven primarily '
        'by strong performance in the enterprise segment. The APAC region contributed '
        '28% of total revenue, up from 21% the previous year.'
    )

    doc.add_paragraph(
        'Customer satisfaction scores improved across all product lines, with the flagship '
        'platform achieving an NPS of 67, the highest recorded in company history. '
        'Key drivers include the new user interface launched in March and expanded '
        'customer success program.'
    )

    doc.add_heading('Key Metrics', level=2)
    # Add a simple table with metrics
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Metric', 'Q3 2024', 'Q4 2024']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    metrics_data = [
        ['Monthly Active Users', '1,240,000', '1,380,000'],
        ['Revenue (USD)', '$4,230,000', '$4,890,000'],
        ['Support Tickets Resolved', '98.2%', '99.1%'],
        ['New Enterprise Clients', '47', '63'],
    ]
    for i, row_data in enumerate(metrics_data, 1):
        for j, val in enumerate(row_data):
            table.cell(i, j).text = val

    doc.add_paragraph()  # spacing

    doc.add_heading('Recommendations', level=2)
    doc.add_paragraph(
        'Based on the analysis, the following strategic initiatives are recommended for 2025:'
    )
    doc.add_paragraph('Expand APAC sales team by 15 headcount', style='List Bullet')
    doc.add_paragraph('Launch mobile-first product variant in H1 2025', style='List Bullet')
    doc.add_paragraph('Invest in AI-driven analytics dashboard', style='List Bullet')
    doc.add_paragraph('Partner with regional resellers in Southeast Asia', style='List Bullet')

    doc.add_paragraph(
        'Implementation of these recommendations is expected to drive 18-22% revenue '
        'growth in FY2025, positioning the company for a potential Series C funding round.'
    )

    # Ensure Desktop directory exists on VM (handled via script logic)
    # Save document
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
