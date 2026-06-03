"""
Initial Setup: Policy update memo with 'effective immediately' appearing 3 times, no highlighting.
Task ID: writer_hr_021
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_021'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_run(para, text, bold=False, font_name="Calibri", font_size=Pt(11), color=None):
    """Helper to add a formatted run."""
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def create_initial():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- MEMO HEADER ---
    heading = doc.add_heading('MEMORANDUM', level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Memo metadata
    meta_lines = [
        ("TO:", "All Employees"),
        ("FROM:", "Rachel Torres, Vice President of Human Resources"),
        ("DATE:", "March 28, 2026"),
        ("RE:", "Updated Workplace Policies and Benefits Changes"),
    ]
    for label, value in meta_lines:
        para = doc.add_paragraph()
        add_run(para, label, bold=True)
        add_run(para, f"  {value}")

    # Horizontal rule (via paragraph border is complex; use dashes)
    rule = doc.add_paragraph()
    add_run(rule, "=" * 72, font_size=Pt(8))

    # --- INTRODUCTION ---
    intro = doc.add_paragraph()
    add_run(intro, (
        "Dear Colleagues,\n\n"
        "Following the Board of Directors meeting held on March 25, 2026, "
        "we are pleased to announce several important updates to our workplace "
        "policies. These changes reflect our ongoing commitment to fostering a "
        "supportive and productive work environment for all team members."
    ))

    # --- SECTION 1 (contains first 'effective immediately') ---
    doc.add_heading('1. Remote Work Policy Update', level=1)

    p1 = doc.add_paragraph()
    add_run(p1, (
        "The company is expanding its flexible work arrangements. Employees in "
        "departments with eligible roles may now request up to three remote work "
        "days per week, subject to manager approval. This updated remote work policy is "
    ))
    add_run(p1, "effective immediately")
    add_run(p1, (
        " and replaces the previous two-day limit established in January 2025. "
        "Managers should review team schedules and ensure adequate office coverage "
        "on all business days."
    ))

    p1a = doc.add_paragraph()
    add_run(p1a, (
        "Please note that certain client-facing roles and laboratory positions will "
        "continue to require on-site presence as determined by department heads. "
        "Requests for exceptions should be submitted through the HR portal."
    ))

    # --- SECTION 2 (contains second 'effective immediately') ---
    doc.add_heading('2. Health and Wellness Benefits Enhancement', level=1)

    p2 = doc.add_paragraph()
    add_run(p2, (
        "We are pleased to announce an enhancement to our employee wellness program. "
        "The annual wellness stipend will increase from $500 to $750 per employee. "
        "Additionally, mental health counseling sessions covered under our Employee "
        "Assistance Program (EAP) will increase from 6 to 12 sessions per year. "
        "These benefit changes are "
    ))
    add_run(p2, "effective immediately")
    add_run(p2, (
        " and will be reflected in your benefits portal within the next five business "
        "days. Employees who have already exhausted their current session allocation "
        "may schedule additional appointments starting April 1, 2026."
    ))

    # --- SECTION 3 (contains third 'effective immediately') ---
    doc.add_heading('3. Updated Code of Conduct — Electronic Communications', level=1)

    p3 = doc.add_paragraph()
    add_run(p3, (
        "In response to evolving cybersecurity threats and data privacy regulations, "
        "the Information Security team has revised Section 4.7 of our Code of Conduct "
        "regarding the use of personal devices and electronic communications. Key updates "
        "include mandatory two-factor authentication for all company systems and a "
        "prohibition on using unauthorized cloud storage services for company data. "
        "The revised code of conduct is "
    ))
    add_run(p3, "effective immediately")
    add_run(p3, (
        ". All employees must complete the updated compliance training module in the "
        "Learning Management System by April 15, 2026. Failure to complete the training "
        "by the deadline may result in temporary suspension of system access."
    ))

    # --- SECTION 4 (no 'effective immediately') ---
    doc.add_heading('4. Office Renovation Schedule', level=1)

    p4 = doc.add_paragraph()
    add_run(p4, (
        "As part of our facilities improvement plan, the third floor conference rooms "
        "and collaborative workspaces will undergo renovation beginning April 14, 2026. "
        "During this period, affected teams will be relocated to temporary workspaces on "
        "the fifth floor. The renovation is expected to be completed by June 30, 2026. "
        "Detailed relocation maps and desk assignments will be distributed by April 7."
    ))

    # --- CLOSING ---
    doc.add_heading('Questions and Contact Information', level=1)

    closing = doc.add_paragraph()
    add_run(closing, (
        "If you have any questions about these policy changes, please do not hesitate to "
        "contact the HR department at hr@globaltech-solutions.com or extension 4200. "
        "You may also schedule a one-on-one meeting with your HR Business Partner "
        "through the company intranet.\n\n"
        "We appreciate your continued dedication and look forward to these improvements "
        "enhancing your work experience.\n\n"
        "Warm regards,"
    ))

    sign = doc.add_paragraph()
    add_run(sign, "Rachel Torres", bold=True)
    sign2 = doc.add_paragraph()
    add_run(sign2, "Vice President of Human Resources")
    sign3 = doc.add_paragraph()
    add_run(sign3, "GlobalTech Solutions, Inc.")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
