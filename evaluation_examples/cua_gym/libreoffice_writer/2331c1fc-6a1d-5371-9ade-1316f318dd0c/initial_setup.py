"""
Initial Setup: Writer document with 5 chapters (Heading 1), no TOC
Task ID: writer_rd_081
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_081'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Chapter 1: Introduction ---
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'The rapid adoption of machine learning techniques across diverse industries '
        'has fundamentally transformed how organizations approach complex decision-making '
        'processes. From healthcare diagnostics to financial risk assessment, automated '
        'systems now play a critical role in augmenting human expertise and improving '
        'operational efficiency.'
    )
    doc.add_paragraph(
        'This report examines the deployment of predictive analytics models within '
        'the context of urban transportation planning. Our study, conducted over a '
        '14-month period from January 2024 to February 2025, analyzed traffic flow '
        'patterns across 37 metropolitan intersections in the greater Portland area. '
        'The findings presented here aim to inform both policy decisions and future '
        'research directions in smart city infrastructure.'
    )
    doc.add_paragraph(
        'Specifically, we investigate three key research questions: (1) Can real-time '
        'sensor data improve traffic signal optimization by more than 15%? (2) What is '
        'the marginal cost-benefit ratio of deploying IoT devices at secondary '
        'intersections? (3) How do seasonal weather patterns interact with commuter '
        'behavior to affect model accuracy?'
    )

    # --- Chapter 2: Methodology ---
    doc.add_heading('Methodology', level=1)
    doc.add_paragraph(
        'Our research methodology combined quantitative traffic analysis with '
        'qualitative interviews of city transportation planners. We deployed 128 '
        'inductive loop sensors and 64 camera-based detection units across the study '
        'area. Data collection occurred continuously, with readings captured at '
        '5-second intervals and aggregated into 15-minute bins for model training.'
    )
    doc.add_paragraph(
        'The primary analytical framework utilized a gradient-boosted decision tree '
        'ensemble (XGBoost v1.7.6) trained on historical traffic volume data from '
        '2020 to 2023. Feature engineering incorporated weather data from NOAA '
        'stations, event calendars from the Portland Metro Events Commission, and '
        'real-time transit ridership data from TriMet. Cross-validation was performed '
        'using a temporal split strategy to prevent data leakage.'
    )
    doc.add_paragraph(
        'Statistical significance was assessed using the Wilcoxon signed-rank test '
        'with a Bonferroni correction for multiple comparisons. Effect sizes were '
        'reported using Cohen\'s d, with thresholds of 0.2 (small), 0.5 (medium), '
        'and 0.8 (large) following standard conventions.'
    )

    # --- Chapter 3: Results ---
    doc.add_heading('Results', level=1)
    doc.add_paragraph(
        'The optimized signal timing model achieved a mean reduction in vehicle '
        'delay of 22.4% (95% CI: 19.1%\u201325.7%) across all monitored intersections '
        'during peak hours (7:00\u20139:00 AM and 4:30\u20136:30 PM). This improvement '
        'exceeded our initial target of 15%, with particularly strong gains observed '
        'at intersections with high left-turn volumes.'
    )
    doc.add_paragraph(
        'Secondary intersections equipped with IoT sensors showed a 14.8% improvement '
        'in throughput compared to the control group. The cost-benefit analysis revealed '
        'a break-even point of 18 months for sensor deployment, with projected annual '
        'savings of $2.3 million in reduced fuel consumption and commuter time across '
        'the study region.'
    )
    doc.add_paragraph(
        'Seasonal analysis revealed that model accuracy degraded by approximately '
        '8.3 percentage points during winter months (November through February), '
        'primarily due to weather-induced variations in driving behavior. Incorporating '
        'real-time precipitation data improved winter predictions by 5.1 percentage '
        'points, partially offsetting this seasonal effect.'
    )

    # --- Chapter 4: Discussion ---
    doc.add_heading('Discussion', level=1)
    doc.add_paragraph(
        'The results demonstrate that machine learning-based signal optimization '
        'can yield substantial improvements in urban traffic flow, surpassing '
        'traditional fixed-timing and actuated control methods. The 22.4% reduction '
        'in vehicle delay aligns with findings from similar studies in Seattle and '
        'Austin, though our results suggest stronger effects at intersections with '
        'complex turning movements.'
    )
    doc.add_paragraph(
        'A notable limitation of our approach is the reliance on sensor infrastructure '
        'that requires ongoing maintenance. During the study period, 11 of 128 loop '
        'sensors experienced failures requiring repair, with an average downtime of '
        '4.2 days. Future implementations should consider redundant sensor arrays '
        'or camera-only solutions to improve system resilience.'
    )
    doc.add_paragraph(
        'The seasonal degradation in model performance highlights the need for '
        'adaptive retraining pipelines. We recommend quarterly model updates with '
        'expanding training windows, supplemented by online learning modules that '
        'can adjust to abrupt weather events within a 2-hour response window.'
    )

    # --- Chapter 5: Conclusion ---
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'This study provides compelling evidence that predictive analytics can '
        'meaningfully improve urban traffic management. The 22.4% reduction in '
        'vehicle delay, combined with favorable cost-benefit ratios for IoT sensor '
        'deployment, supports the expansion of smart infrastructure programs in '
        'mid-sized metropolitan areas.'
    )
    doc.add_paragraph(
        'Future work should explore the integration of connected vehicle data as '
        'an additional input source, which could further improve prediction accuracy '
        'and reduce dependence on fixed-location sensors. Additionally, the ethical '
        'implications of algorithmic traffic management\u2014including equitable access '
        'across neighborhoods with varying socioeconomic profiles\u2014merit careful '
        'examination as these systems scale.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
