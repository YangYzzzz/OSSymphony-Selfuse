"""
Reward Script: Horizontal TOC navigation header in Writer document
Task ID: writer_rd_081
Domain: libreoffice_writer
Scoring:
  C1 (0.30) — TOC line at top contains all 5 heading names separated by ' | '
  C2 (0.20) — TOC paragraph is center-aligned and 11pt font
  C3 (0.20) — TOC entries are hyperlinks (anchor links to sections)
  C4 (0.15) — Horizontal line (bottom border) separates TOC from body
  C5 (0.15) — Original body content preserved (headings still present after TOC)
"""

import os
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_081'

EXPECTED_HEADINGS = ['Introduction', 'Methodology', 'Results', 'Discussion', 'Conclusion']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # ─── Component 1: TOC line at top with all 5 headings separated by ' | ' (0.30 pts) ───
    try:
        # The first paragraph should be the TOC navigation line
        p0 = doc.paragraphs[0]
        p0_text = p0.text.strip()

        # Check that it contains all expected headings separated by |
        has_all_headings = True
        for heading in EXPECTED_HEADINGS:
            if heading not in p0_text:
                has_all_headings = False
                print(f"FAIL: Component 1 — heading '{heading}' not found in TOC line")
                break

        has_separator = ' | ' in p0_text
        # The TOC should NOT be a Heading style — it's a navigation line
        is_not_heading_style = 'Heading' not in p0.style.name

        if has_all_headings and has_separator and is_not_heading_style:
            # Verify exact order
            parts = [part.strip() for part in p0_text.split('|')]
            parts_clean = [p for p in parts if p]
            if parts_clean == EXPECTED_HEADINGS:
                print(f"PASS: Component 1 — TOC line at top with correct headings in order (0.30 pts)")
                total_score += 0.30
            elif set(parts_clean) == set(EXPECTED_HEADINGS):
                # Headings present but wrong order — partial
                print(f"PASS (partial): Component 1 — headings present but wrong order (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 1 — TOC parts {parts_clean} don't match expected {EXPECTED_HEADINGS}")
        else:
            if not has_all_headings:
                pass  # already printed
            if not has_separator:
                print(f"FAIL: Component 1 — no ' | ' separator found in first paragraph: {repr(p0_text[:100])}")
            if not is_not_heading_style:
                print(f"FAIL: Component 1 — first paragraph is Heading style, not a TOC nav line")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ─── Component 2: TOC paragraph is center-aligned and 11pt font (0.20 pts) ───
    try:
        p0 = doc.paragraphs[0]
        alignment = p0.paragraph_format.alignment
        is_centered = alignment == WD_PARAGRAPH_ALIGNMENT.CENTER

        # Check font size — runs should be 11pt (w:sz val="22" = 11pt)
        run_sizes = []
        for run in p0.runs:
            if run.font.size is not None:
                run_sizes.append(run.font.size.pt)

        # Also check hyperlink runs via XML (para.runs may miss hyperlink text)
        sz_elements = p0._element.findall('.//w:sz', ns)
        xml_sizes = set()
        for sz_el in sz_elements:
            val = sz_el.get(qn('w:val'))
            if val:
                xml_sizes.add(int(val) / 2.0)  # half-points to points

        is_11pt = False
        if run_sizes and all(abs(s - 11.0) < 0.5 for s in run_sizes):
            is_11pt = True
        elif xml_sizes and all(abs(s - 11.0) < 0.5 for s in xml_sizes):
            is_11pt = True

        if is_centered and is_11pt:
            print(f"PASS: Component 2 — center-aligned and 11pt font (0.20 pts)")
            total_score += 0.20
        elif is_centered:
            print(f"PASS (partial): Component 2 — center-aligned but font size issue (0.10 pts)")
            print(f"  Run sizes: {run_sizes}, XML sizes: {xml_sizes}")
            total_score += 0.10
        elif is_11pt:
            print(f"PASS (partial): Component 2 — 11pt font but not centered (0.10 pts)")
            print(f"  Alignment: {alignment}")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — not centered (align={alignment}) and font size issue")
            print(f"  Run sizes: {run_sizes}, XML sizes: {xml_sizes}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ─── Component 3: TOC entries are hyperlinks (0.20 pts) ───
    try:
        p0 = doc.paragraphs[0]
        hyperlinks = p0._element.findall('.//w:hyperlink', ns)
        hyperlink_count = len(hyperlinks)

        # Extract hyperlink texts
        hl_texts = []
        for h in hyperlinks:
            texts = [t.text for t in h.findall('.//w:t', ns) if t.text]
            hl_texts.append(''.join(texts).strip())

        if hyperlink_count >= 5:
            # Check that all 5 headings are hyperlinked
            matched = sum(1 for heading in EXPECTED_HEADINGS if heading in hl_texts)
            if matched == 5:
                print(f"PASS: Component 3 — all 5 headings are hyperlinks (0.20 pts)")
                total_score += 0.20
            else:
                ratio = matched / 5.0
                pts = round(0.20 * ratio, 2)
                print(f"PASS (partial): Component 3 — {matched}/5 headings hyperlinked ({pts} pts)")
                total_score += pts
        elif hyperlink_count > 0:
            ratio = min(hyperlink_count / 5.0, 1.0)
            pts = round(0.20 * ratio, 2)
            print(f"PASS (partial): Component 3 — {hyperlink_count} hyperlinks found ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 3 — no hyperlinks found in TOC paragraph")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ─── Component 4: Horizontal line / border separating TOC from body (0.15 pts) ───
    try:
        # The horizontal line is implemented as a bottom border on a paragraph
        # near the top (P0 or P1). Check the first few paragraphs for a bottom border.
        border_found = False
        for idx in range(min(3, len(doc.paragraphs))):
            p = doc.paragraphs[idx]
            ppr = p._element.find(qn('w:pPr'))
            if ppr is not None:
                pbdr = ppr.find(qn('w:pBdr'))
                if pbdr is not None:
                    bottom = pbdr.find(qn('w:bottom'))
                    if bottom is not None:
                        border_val = bottom.get(qn('w:val'))
                        if border_val and border_val != 'none':
                            border_found = True
                            print(f"PASS: Component 4 — horizontal line found (bottom border on P{idx}) (0.15 pts)")
                            total_score += 0.15
                            break

        if not border_found:
            print(f"FAIL: Component 4 — no horizontal line (bottom border) found in first 3 paragraphs")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ─── Component 5: TOC inserted AND original body content preserved (0.15 pts) ───
    # This is a compound check: the document must have the new TOC paragraphs
    # at the top AND still retain all 5 original Heading 1 paragraphs.
    # The initial file has 19 paragraphs starting with Heading 1.
    # The golden has 21+ paragraphs with a non-Heading first paragraph.
    try:
        # Gate: first paragraph must NOT be Heading 1 (proves TOC was inserted)
        first_is_toc = (doc.paragraphs[0].style.name != 'Heading 1')

        heading1_paras = []
        for p in doc.paragraphs:
            if p.style and p.style.name == 'Heading 1':
                heading1_paras.append(p.text.strip())

        matched_headings = sum(1 for h in EXPECTED_HEADINGS if h in heading1_paras)

        if first_is_toc and matched_headings == 5:
            print(f"PASS: Component 5 — TOC inserted at top AND all 5 headings preserved (0.15 pts)")
            total_score += 0.15
        elif first_is_toc and matched_headings >= 3:
            pts = round(0.15 * (matched_headings / 5.0), 2)
            print(f"PASS (partial): Component 5 — TOC at top, {matched_headings}/5 headings ({pts} pts)")
            total_score += pts
        else:
            if not first_is_toc:
                print(f"FAIL: Component 5 — first paragraph is still Heading 1, no TOC inserted")
            else:
                print(f"FAIL: Component 5 — only {matched_headings}/5 Heading 1 paragraphs found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
