"""
Initial Setup: Position image in Writer document
Task ID: writer_obj_027
Domain: libreoffice_writer

Creates 'precise_doc.docx' with a floating image (5cm x 4cm) placed roughly in
the middle of page 1 (X: 7cm, Y: 10cm from margin). The task is to reposition
it to exactly X: 2cm, Y: 5cm from the margins.
"""

import os
import io
import shlex
import subprocess
import time
from lxml import etree
from PIL import Image
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_027'
OUTPUT = f'{WORKDIR}/Desktop/precise_doc.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_image_bytes():
    """Create a simple professional-looking PNG image (300x240 px)."""
    img = Image.new('RGB', (300, 240), color=(220, 230, 242))
    # Add a simple border region
    from PIL import ImageDraw
    draw = ImageDraw.Draw(img)
    draw.rectangle([5, 5, 294, 234], outline=(70, 114, 184), width=3)
    draw.rectangle([20, 20, 280, 220], fill=(240, 245, 255))
    # Simple diagram bars
    draw.rectangle([50, 60, 80, 180], fill=(70, 114, 184))
    draw.rectangle([110, 90, 140, 180], fill=(112, 173, 71))
    draw.rectangle([170, 50, 200, 180], fill=(255, 192, 0))
    draw.rectangle([230, 110, 260, 180], fill=(192, 0, 0))
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes


def add_floating_image(doc, img_path, x_cm, y_cm, width_cm, height_cm,
                        doc_pr_id=1, img_name='Picture 1'):
    """
    Add a floating image to the document at an absolute position.
    x_cm, y_cm: position from the page margin (in cm)
    width_cm, height_cm: image dimensions (in cm)
    Uses wp:anchor with relativeFrom='margin' for both H and V.
    """
    x_emu = int(Cm(x_cm))
    y_emu = int(Cm(y_cm))
    cx = int(Cm(width_cm))
    cy = int(Cm(height_cm))

    r_id, _ = doc.part.get_or_add_image(img_path)

    WP  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    A   = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    PIC = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

    def qnt(ns, tag):
        return '{%s}%s' % (ns, tag)

    anchor = etree.Element(qnt(WP, 'anchor'))
    anchor.set('distT', '0'); anchor.set('distB', '0')
    anchor.set('distL', '114300'); anchor.set('distR', '114300')
    anchor.set('simplePos', '0'); anchor.set('relativeHeight', '251658240')
    anchor.set('behindDoc', '0'); anchor.set('locked', '0')
    anchor.set('layoutInCell', '1'); anchor.set('allowOverlap', '1')

    simplePos = etree.SubElement(anchor, qnt(WP, 'simplePos'))
    simplePos.set('x', '0'); simplePos.set('y', '0')

    posH = etree.SubElement(anchor, qnt(WP, 'positionH'))
    posH.set('relativeFrom', 'margin')
    posOffset_h = etree.SubElement(posH, qnt(WP, 'posOffset'))
    posOffset_h.text = str(x_emu)

    posV = etree.SubElement(anchor, qnt(WP, 'positionV'))
    posV.set('relativeFrom', 'margin')
    posOffset_v = etree.SubElement(posV, qnt(WP, 'posOffset'))
    posOffset_v.text = str(y_emu)

    extent = etree.SubElement(anchor, qnt(WP, 'extent'))
    extent.set('cx', str(cx)); extent.set('cy', str(cy))

    effectExtent = etree.SubElement(anchor, qnt(WP, 'effectExtent'))
    effectExtent.set('l', '0'); effectExtent.set('t', '0')
    effectExtent.set('r', '0'); effectExtent.set('b', '0')

    wrapSquare = etree.SubElement(anchor, qnt(WP, 'wrapSquare'))
    wrapSquare.set('wrapText', 'bothSides')

    docPr = etree.SubElement(anchor, qnt(WP, 'docPr'))
    docPr.set('id', str(doc_pr_id)); docPr.set('name', img_name)

    cNvGraphicFramePr = etree.SubElement(anchor, qnt(WP, 'cNvGraphicFramePr'))
    graphicFrameLocks = etree.SubElement(cNvGraphicFramePr, qnt(A, 'graphicFrameLocks'))
    graphicFrameLocks.set('noChangeAspect', '1')

    graphic = etree.SubElement(anchor, qnt(A, 'graphic'))
    graphicData = etree.SubElement(graphic, qnt(A, 'graphicData'))
    graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')

    pic_elem = etree.SubElement(graphicData, qnt(PIC, 'pic'))
    nvPicPr = etree.SubElement(pic_elem, qnt(PIC, 'nvPicPr'))
    cNvPr = etree.SubElement(nvPicPr, qnt(PIC, 'cNvPr'))
    cNvPr.set('id', '0'); cNvPr.set('name', 'chart_diagram.png')
    cNvPicPr = etree.SubElement(nvPicPr, qnt(PIC, 'cNvPicPr'))

    blipFill = etree.SubElement(pic_elem, qnt(PIC, 'blipFill'))
    blip = etree.SubElement(blipFill, qnt(A, 'blip'))
    blip.set(qnt(R_NS, 'embed'), r_id)
    stretch = etree.SubElement(blipFill, qnt(A, 'stretch'))
    etree.SubElement(stretch, qnt(A, 'fillRect'))

    spPr = etree.SubElement(pic_elem, qnt(PIC, 'spPr'))
    xfrm = etree.SubElement(spPr, qnt(A, 'xfrm'))
    off = etree.SubElement(xfrm, qnt(A, 'off'))
    off.set('x', '0'); off.set('y', '0')
    ext = etree.SubElement(xfrm, qnt(A, 'ext'))
    ext.set('cx', str(cx)); ext.set('cy', str(cy))
    prstGeom = etree.SubElement(spPr, qnt(A, 'prstGeom'))
    prstGeom.set('prst', 'rect')
    etree.SubElement(prstGeom, qnt(A, 'avLst'))

    drawing = etree.Element(qnt(W, 'drawing'))
    drawing.append(anchor)

    para = doc.add_paragraph()
    run = para.add_run()
    run._element.append(drawing)
    return para


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    # Save the image to a temp file on the VM (we need a file path for get_or_add_image)
    img_tmp = '/tmp/writer_obj_027_chart.png'
    img_bytes = create_image_bytes()
    with open(img_tmp, 'wb') as f:
        f.write(img_bytes.read())

    doc = Document()

    # Page setup: A4 with standard margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # ---- Heading ----
    heading = doc.add_heading('', level=1)
    run_h = heading.add_run('Q1 2025 Performance Review')
    run_h.font.name = 'Calibri'
    run_h.font.size = Pt(16)

    # ---- Introductory paragraph ----
    para1 = doc.add_paragraph()
    run1 = para1.add_run(
        'This report summarizes the key performance indicators for the first quarter '
        'of fiscal year 2025. All metrics have been compiled from departmental submissions '
        'reviewed by Finance and Operations leadership.'
    )
    run1.font.name = 'Calibri'
    run1.font.size = Pt(11)

    # ---- Section heading ----
    para2 = doc.add_paragraph()
    run2 = para2.add_run('Revenue Analysis')
    run2.font.name = 'Calibri'
    run2.font.size = Pt(13)
    run2.bold = True

    # ---- Body text ----
    para3 = doc.add_paragraph()
    run3 = para3.add_run(
        'Total revenue for Q1 reached $4.72M, representing a 12.4% increase over the '
        'same period last year. North American operations contributed $2.1M while '
        'APAC saw the strongest growth at 18.7% year-over-year.'
    )
    run3.font.name = 'Calibri'
    run3.font.size = Pt(11)

    # ---- Floating image: placed roughly in the middle of the page ----
    # Initial position: X=7cm, Y=10cm from margin (clearly NOT the target 2cm, 5cm)
    add_floating_image(
        doc,
        img_path=img_tmp,
        x_cm=7.0,   # NOT the target: 2cm
        y_cm=10.0,  # NOT the target: 5cm
        width_cm=5.0,
        height_cm=4.0,
        doc_pr_id=1,
        img_name='Performance Chart'
    )

    # ---- More body text ----
    para4 = doc.add_paragraph()
    run4 = para4.add_run(
        'The chart above illustrates revenue by business unit across the four main '
        'product lines. Engineering Services and Cloud Solutions together account for '
        '68% of total revenue.'
    )
    run4.font.name = 'Calibri'
    run4.font.size = Pt(11)

    para5 = doc.add_paragraph()
    run5 = para5.add_run('Key Highlights')
    run5.font.name = 'Calibri'
    run5.font.size = Pt(13)
    run5.bold = True

    bullets = [
        'Engineering Services: $1.82M (+9.2% YoY)',
        'Cloud Solutions: $1.39M (+22.1% YoY)',
        'Professional Services: $0.98M (+5.6% YoY)',
        'Licensing & Support: $0.53M (+3.1% YoY)',
    ]
    for bullet in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        br = bp.add_run(bullet)
        br.font.name = 'Calibri'
        br.font.size = Pt(11)

    para6 = doc.add_paragraph()
    run6 = para6.add_run(
        'Operating expenses remained within budget at $3.41M against a projected $3.55M. '
        'Headcount grew by 14 full-time equivalents across Engineering and Sales teams.'
    )
    run6.font.name = 'Calibri'
    run6.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Clean up temp image
    os.remove(img_tmp)

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
