"""
Initial Setup: ML Workflow Specification Document
Task ID: writer_struct_061
Domain: libreoffice_writer
Creates an 8-page technical specification for ML workflow with empty document properties.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'ml_spec'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Page 1: Title Page ---
    title_para = doc.add_heading('Machine Learning Workflow Specification', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph('Technical Reference Document')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True

    version_para = doc.add_paragraph('Version 2.1 | Confidential')
    version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    version_para.runs[0].font.size = Pt(11)

    doc.add_paragraph('')
    date_para = doc.add_paragraph('Date: March 2025')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    dept_para = doc.add_paragraph('AI Engineering Team — Research Division')
    dept_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # --- Page 2: Table of Contents ---
    doc.add_heading('Table of Contents', level=1)

    toc_items = [
        ('1. Introduction', '3'),
        ('2. System Architecture', '3'),
        ('3. Data Pipeline Components', '4'),
        ('4. Model Training Workflow', '5'),
        ('5. Evaluation Framework', '6'),
        ('6. Deployment Pipeline', '7'),
        ('7. Monitoring and Alerting', '7'),
        ('8. Appendix', '8'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run(f'\t{page}')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

    # --- Page 3: Introduction ---
    doc.add_heading('1. Introduction', level=1)

    doc.add_paragraph(
        'This document specifies the complete technical architecture and workflow for the Machine Learning '
        'pipeline developed by the AI Engineering Team. The specification covers data ingestion, preprocessing, '
        'model training, evaluation, and deployment procedures for production ML systems.'
    )

    doc.add_paragraph(
        'The ML Workflow Specification serves as the authoritative reference for all engineering teams involved '
        'in the design, development, and maintenance of the organization\'s machine learning infrastructure. '
        'All systems must conform to the standards and protocols defined herein.'
    )

    doc.add_heading('1.1 Scope', level=2)
    doc.add_paragraph(
        'This specification applies to all machine learning projects within the organization that involve: '
        'supervised learning tasks, unsupervised learning experiments, reinforcement learning environments, '
        'and hybrid approaches combining multiple paradigms. The document does not cover rule-based systems, '
        'traditional statistical models without learned parameters, or data warehousing infrastructure.'
    )

    doc.add_heading('1.2 Document Conventions', level=2)
    doc.add_paragraph(
        'Throughout this document, the following conventions are used: MUST indicates a mandatory requirement, '
        'SHOULD indicates a recommended practice, and MAY indicates an optional capability. '
        'Code snippets are rendered in monospace font. Configuration parameters are shown in angle brackets.'
    )

    doc.add_page_break()

    # --- Page 4: System Architecture ---
    doc.add_heading('2. System Architecture', level=1)

    doc.add_paragraph(
        'The ML platform is built on a microservices architecture designed for horizontal scalability '
        'and fault tolerance. Each component is containerized and orchestrated via Kubernetes, enabling '
        'independent scaling of compute-intensive pipeline stages.'
    )

    doc.add_heading('2.1 Core Components', level=2)

    components_table = doc.add_table(rows=6, cols=3)
    components_table.style = 'Table Grid'

    # Header row
    headers = ['Component', 'Technology', 'Description']
    for j, h in enumerate(headers):
        cell = components_table.cell(0, j)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    component_data = [
        ('Data Ingestion Service', 'Apache Kafka 3.5', 'Real-time streaming ingestion from multiple sources'),
        ('Feature Store', 'Redis + PostgreSQL', 'Centralized feature computation and serving layer'),
        ('Training Orchestrator', 'MLflow + Kubeflow', 'Manages distributed model training jobs'),
        ('Model Registry', 'MLflow Registry', 'Versioned model storage with metadata'),
        ('Serving Layer', 'TorchServe / TF-Serving', 'Low-latency model inference endpoints'),
    ]
    for i, row_data in enumerate(component_data, 1):
        for j, val in enumerate(row_data):
            components_table.cell(i, j).text = val

    doc.add_paragraph('')

    doc.add_heading('2.2 Infrastructure Requirements', level=2)
    doc.add_paragraph(
        'Production deployments require a minimum cluster configuration of 8 CPU nodes (32 cores each), '
        '4 GPU nodes (NVIDIA A100 80GB), 500TB NVMe storage for feature store, and 10Gbps network fabric '
        'with RDMA support for distributed training workloads.'
    )

    doc.add_page_break()

    # --- Page 5: Data Pipeline ---
    doc.add_heading('3. Data Pipeline Components', level=1)

    doc.add_paragraph(
        'The data pipeline is the foundational layer of the ML workflow, responsible for acquiring raw data, '
        'applying transformations, and delivering clean, feature-engineered datasets to downstream consumers. '
        'Pipeline reliability is critical; SLA requires 99.9% uptime and end-to-end latency under 500ms.'
    )

    doc.add_heading('3.1 Data Sources', level=2)

    sources = [
        'Transactional databases (PostgreSQL, MySQL) via CDC connectors',
        'Event streams from user interaction tracking systems',
        'Third-party API feeds (weather, market data, social signals)',
        'Batch uploads from partner organizations (CSV, Parquet formats)',
        'IoT sensor streams via MQTT brokers',
    ]
    for src in sources:
        doc.add_paragraph(src, style='List Bullet')

    doc.add_heading('3.2 Preprocessing Steps', level=2)

    prep_table = doc.add_table(rows=7, cols=3)
    prep_table.style = 'Table Grid'

    prep_headers = ['Step', 'Operation', 'Config Parameter']
    for j, h in enumerate(prep_headers):
        cell = prep_table.cell(0, j)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    prep_data = [
        ('1', 'Schema validation', '<schema_registry_url>'),
        ('2', 'Null imputation', '<imputation_strategy>'),
        ('3', 'Outlier detection', '<zscore_threshold>'),
        ('4', 'Normalization', '<scaler_type>'),
        ('5', 'Feature encoding', '<encoding_config>'),
        ('6', 'Train/val/test split', '<split_ratios>'),
    ]
    for i, row_data in enumerate(prep_data, 1):
        for j, val in enumerate(row_data):
            prep_table.cell(i, j).text = val

    doc.add_page_break()

    # --- Page 6: Model Training ---
    doc.add_heading('4. Model Training Workflow', level=1)

    doc.add_paragraph(
        'Model training follows a standardized workflow to ensure reproducibility and comparability across '
        'experiments. All training runs MUST be tracked in the experiment registry with full parameter '
        'logging, artifact storage, and performance metrics.'
    )

    doc.add_heading('4.1 Training Configuration', level=2)
    doc.add_paragraph(
        'Each training job is defined by a YAML configuration file specifying model architecture, '
        'hyperparameters, data sources, and compute requirements. Configuration versioning follows '
        'semantic versioning (MAJOR.MINOR.PATCH).'
    )

    doc.add_heading('4.2 Distributed Training', level=2)
    doc.add_paragraph(
        'Large-scale models exceeding 1B parameters MUST use distributed training with data parallelism '
        'across at least 4 GPU nodes. The system supports PyTorch DDP (DistributedDataParallel) and '
        'DeepSpeed ZeRO optimization stages 1, 2, and 3.'
    )

    doc.add_heading('4.3 Hyperparameter Optimization', level=2)
    doc.add_paragraph(
        'Hyperparameter search uses Optuna with a TPE sampler. Each HPO study runs a minimum of 50 trials '
        'with early stopping via median pruner. Studies are persisted in the PostgreSQL backend for resumability.'
    )

    training_params = doc.add_table(rows=6, cols=2)
    training_params.style = 'Table Grid'
    param_headers = ['Parameter', 'Default Value']
    for j, h in enumerate(param_headers):
        cell = training_params.cell(0, j)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    param_data = [
        ('learning_rate', '3e-4'),
        ('batch_size', '256'),
        ('max_epochs', '100'),
        ('early_stopping_patience', '10'),
        ('gradient_clip_value', '1.0'),
    ]
    for i, row_data in enumerate(param_data, 1):
        for j, val in enumerate(row_data):
            training_params.cell(i, j).text = val

    doc.add_page_break()

    # --- Page 7: Evaluation Framework ---
    doc.add_heading('5. Evaluation Framework', level=1)

    doc.add_paragraph(
        'All models are subject to a multi-stage evaluation protocol before advancing to the next pipeline stage. '
        'Evaluation criteria include task-specific performance metrics, computational efficiency benchmarks, '
        'fairness assessments, and adversarial robustness tests.'
    )

    doc.add_heading('5.1 Performance Metrics', level=2)

    metrics = [
        'Classification: Accuracy, F1-score (macro/micro), ROC-AUC, PR-AUC',
        'Regression: RMSE, MAE, MAPE, R² coefficient of determination',
        'Ranking: NDCG@K, MRR, MAP',
        'Generation: BLEU-4, ROUGE-L, BERTScore, Perplexity',
        'Detection: mAP@0.5, mAP@0.5:0.95, FPS throughput',
    ]
    for m in metrics:
        doc.add_paragraph(m, style='List Bullet')

    doc.add_heading('5.2 Minimum Acceptance Thresholds', level=2)
    doc.add_paragraph(
        'Models must meet minimum acceptance thresholds before deployment approval. '
        'Thresholds are defined per-domain and reviewed quarterly by the model governance committee. '
        'Any model falling below thresholds requires a root cause analysis and remediation plan.'
    )

    doc.add_heading('6. Deployment Pipeline', level=1)

    doc.add_paragraph(
        'Approved models are deployed via a CI/CD pipeline that automates containerization, integration testing, '
        'canary deployment, and full rollout. Blue-green deployment is used for zero-downtime updates.'
    )

    doc.add_heading('6.1 Canary Deployment', level=2)
    doc.add_paragraph(
        'New model versions receive 5% of production traffic during canary phase. '
        'Canary phase lasts 24-72 hours depending on traffic volume. '
        'Automatic rollback triggers if error rate exceeds 0.1% or p99 latency exceeds 200ms.'
    )

    doc.add_page_break()

    # --- Page 8: Monitoring and Appendix ---
    doc.add_heading('7. Monitoring and Alerting', level=1)

    doc.add_paragraph(
        'Production ML systems are monitored via a unified observability stack (Prometheus + Grafana + Jaeger). '
        'Key monitoring dimensions include model performance drift, data distribution shift, '
        'infrastructure health, and business KPI impact.'
    )

    doc.add_heading('7.1 Drift Detection', level=2)
    doc.add_paragraph(
        'Statistical drift is detected using the Kolmogorov-Smirnov test for continuous features '
        'and chi-squared test for categorical features. Alerts fire when drift exceeds the configured '
        'threshold (default p-value < 0.05). Weekly reference data windows are maintained for comparison.'
    )

    doc.add_heading('7.2 Alert Escalation Policy', level=2)
    alert_items = [
        'P0: Model accuracy drop > 10% — Page on-call ML engineer immediately',
        'P1: Data pipeline SLA breach — Alert team Slack channel within 5 minutes',
        'P2: Feature store latency > 50ms p99 — Ticket created, response within 4 hours',
        'P3: Minor metric degradation < 2% — Weekly review cycle',
    ]
    for a in alert_items:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_heading('8. Appendix', level=1)

    doc.add_heading('8.1 Glossary', level=2)
    glossary_table = doc.add_table(rows=6, cols=2)
    glossary_table.style = 'Table Grid'
    gloss_headers = ['Term', 'Definition']
    for j, h in enumerate(gloss_headers):
        cell = glossary_table.cell(0, j)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    glossary_data = [
        ('CDC', 'Change Data Capture — technique for tracking database changes'),
        ('DDP', 'DistributedDataParallel — PyTorch distributed training paradigm'),
        ('HPO', 'Hyperparameter Optimization — automated search over hyperparameter space'),
        ('MLOps', 'Machine Learning Operations — practices for ML in production'),
        ('SLA', 'Service Level Agreement — performance and availability commitments'),
    ]
    for i, row_data in enumerate(glossary_data, 1):
        for j, val in enumerate(row_data):
            glossary_table.cell(i, j).text = val

    doc.add_heading('8.2 References', level=2)
    references = [
        'MLflow Documentation v2.10 — https://mlflow.org/docs/latest/',
        'Kubeflow Pipelines v1.8 — https://www.kubeflow.org/docs/',
        'PyTorch Distributed Training Guide — https://pytorch.org/tutorials/distributed/',
        'Optuna Documentation — https://optuna.readthedocs.io/',
    ]
    for ref in references:
        doc.add_paragraph(ref, style='List Number')

    # NOTE: Document properties (title, subject, author, keywords) are intentionally left EMPTY
    # The task requires the user to set these via File > Properties > Description tab
    # Verify core_properties are empty:
    doc.core_properties.title = ''
    doc.core_properties.subject = ''
    doc.core_properties.author = ''
    doc.core_properties.keywords = ''

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
