"""
Reward Script: Insert endnotes for 8 technical references in Research_Methods.docx
Task ID: writer_pd_044
Domain: libreoffice_writer
Scoring:
  Component 1: 8 actual endnotes exist in endnotes.xml (0.30 pts)
  Component 2: 8 endnote references in document body (0.20 pts)
  Component 3: Bracketed placeholders [i]-[viii] removed (0.15 pts)
  Component 4: Endnote numbering format is lowerRoman (0.15 pts)
  Component 5: 'Endnotes' heading present at end of document (0.20 pts)
"""

import os
import re
import zipfile
import xml.etree.ElementTree as ET

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_044'

# Persistence hook: save any unsaved LibreOffice edits before verification
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Precondition: file must be a valid docx (zip)
    try:
        zf = zipfile.ZipFile(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot open file {file_path} as zip: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: 8 actual endnotes exist in endnotes.xml (0.30 points)
    # Initial has NO endnotes.xml; golden has 10 entries (2 separators + 8 real)
    try:
        if 'word/endnotes.xml' not in zf.namelist():
            print("FAIL: Component 1 -- no endnotes.xml found in document")
        else:
            tree = ET.parse(zf.open('word/endnotes.xml'))
            root = tree.getroot()
            all_endnotes = root.findall('.//w:endnote', ns)
            # Filter out separator/continuationSeparator (type attribute present)
            real_endnotes = [
                en for en in all_endnotes
                if en.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') is None
            ]
            count = len(real_endnotes)
            if count == 8:
                print(f"PASS: Component 1 -- found exactly 8 real endnotes (0.30 pts)")
                total_score += 0.30
            elif count > 0:
                # Partial credit: proportional to how many of the 8 are present
                partial = 0.30 * min(count, 8) / 8
                print(f"PARTIAL: Component 1 -- found {count}/8 endnotes ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 1 -- found 0 real endnotes")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: 8 endnote references in document body (0.20 points)
    # Initial has 0 references; golden has 8
    try:
        tree2 = ET.parse(zf.open('word/document.xml'))
        root2 = tree2.getroot()
        endnote_refs = root2.findall('.//w:endnoteReference', ns)
        ref_count = len(endnote_refs)
        if ref_count == 8:
            print(f"PASS: Component 2 -- found exactly 8 endnote references in body (0.20 pts)")
            total_score += 0.20
        elif ref_count > 0:
            partial = 0.20 * min(ref_count, 8) / 8
            print(f"PARTIAL: Component 2 -- found {ref_count}/8 endnote references ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- found 0 endnote references in body")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Bracketed placeholders [i]-[viii] removed (0.15 points)
    # Initial has 8 placeholders; golden has 0
    try:
        from docx import Document
        doc = Document(file_path)
        placeholder_pattern = re.compile(r'\[(?:i|ii|iii|iv|v|vi|vii|viii)\]')
        placeholder_count = 0
        for para in doc.paragraphs:
            matches = placeholder_pattern.findall(para.text)
            placeholder_count += len(matches)

        if placeholder_count == 0:
            print(f"PASS: Component 3 -- no bracketed placeholders remaining (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 -- {placeholder_count} bracketed placeholders still present")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Endnote numbering format is lowerRoman (0.15 points)
    # Initial has no endnotePr; golden has lowerRoman + docEnd
    try:
        tree_doc = ET.parse(zf.open('word/document.xml'))
        root_doc = tree_doc.getroot()
        # Check section properties for endnotePr
        endnote_prs = root_doc.findall('.//w:endnotePr', ns)

        # Also check settings.xml
        found_lower_roman = False
        for epr in endnote_prs:
            num_fmt = epr.find('w:numFmt', ns)
            if num_fmt is not None:
                fmt_val = num_fmt.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if fmt_val == 'lowerRoman':
                    found_lower_roman = True
                    break

        # Also check settings.xml if not found in document.xml
        if not found_lower_roman and 'word/settings.xml' in zf.namelist():
            tree_settings = ET.parse(zf.open('word/settings.xml'))
            root_settings = tree_settings.getroot()
            settings_epr = root_settings.findall('.//w:endnotePr', ns)
            for epr in settings_epr:
                num_fmt = epr.find('w:numFmt', ns)
                if num_fmt is not None:
                    fmt_val = num_fmt.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if fmt_val == 'lowerRoman':
                        found_lower_roman = True
                        break

        if found_lower_roman:
            print(f"PASS: Component 4 -- endnote numbering is lowerRoman (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 4 -- endnote numbering is not lowerRoman")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    # Component 5: 'Endnotes' heading present at end of document (0.20 points)
    # Initial has no such heading; golden has it as the last paragraph
    try:
        from docx import Document
        doc = Document(file_path)
        # Look for an 'Endnotes' heading in the last 5 paragraphs
        found_endnotes_heading = False
        last_paras = doc.paragraphs[-5:] if len(doc.paragraphs) >= 5 else doc.paragraphs
        for para in last_paras:
            if 'endnotes' in para.text.lower().strip() and 'Heading' in para.style.name:
                found_endnotes_heading = True
                print(f"  Detail: Found heading '{para.text}' with style '{para.style.name}'")
                break

        if found_endnotes_heading:
            print(f"PASS: Component 5 -- 'Endnotes' heading found near end of document (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 5 -- no 'Endnotes' heading found near end of document")
    except Exception as e:
        print(f"ERROR: Component 5 -- {e}")

    zf.close()

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
