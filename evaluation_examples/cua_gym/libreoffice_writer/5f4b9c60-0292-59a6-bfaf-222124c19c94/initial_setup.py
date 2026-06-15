"""
Initial Setup: Insert endnotes for 8 technical references in Research_Methods.docx
Task ID: writer_pd_044
Domain: libreoffice_writer

Creates a 12-page research methods document with 8 bracketed placeholders [i]-[viii]
where endnote references should be inserted. No endnotes or footnotes exist yet.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_044'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DISPLAY_NAME = f'{WORKDIR}/Research_Methods.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard letter size with 1-inch margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ============================================================
    # TITLE PAGE (page 1)
    # ============================================================
    title = doc.add_heading('Research Methods in Computational Linguistics', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('A Comprehensive Survey of Quantitative and Qualitative Approaches')
    run.font.size = Pt(14)
    run.italic = True

    author_block = doc.add_paragraph()
    author_block.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_block.paragraph_format.space_before = Pt(48)
    run = author_block.add_run('Dr. Elena Vasquez\nDepartment of Linguistics\nStanford University\n\nDr. Raj Patel\nSchool of Computer Science\nCarnegie Mellon University')
    run.font.size = Pt(11)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.space_before = Pt(36)
    run = date_para.add_run('March 2026')
    run.font.size = Pt(11)

    doc.add_page_break()

    # ============================================================
    # TABLE OF CONTENTS (page 2)
    # ============================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1. Introduction', '3'),
        ('2. Literature Review', '4'),
        ('3. Quantitative Methods', '5'),
        ('   3.1 Corpus Analysis', '5'),
        ('   3.2 Statistical Modeling', '6'),
        ('4. Qualitative Methods', '7'),
        ('   4.1 Discourse Analysis', '7'),
        ('   4.2 Ethnographic Approaches', '8'),
        ('5. Mixed Methods Design', '9'),
        ('6. Data Collection Procedures', '10'),
        ('7. Validity and Reliability', '11'),
        ('8. Conclusions', '12'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(f'{item}')
        p.add_run(f'  {"." * (50 - len(item))}  {page}')
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ============================================================
    # CHAPTER 1: INTRODUCTION (pages 3-4)
    # ============================================================
    doc.add_heading('1. Introduction', level=1)

    doc.add_paragraph(
        'The field of computational linguistics has undergone a remarkable transformation over the past '
        'two decades. With the emergence of large-scale language models, neural architectures, and '
        'unprecedented volumes of digitized text, researchers face both extraordinary opportunities and '
        'significant methodological challenges. This survey examines the principal research methods '
        'employed in contemporary computational linguistics, drawing from studies conducted between '
        '2018 and 2025.'
    )

    doc.add_paragraph(
        'The motivation for this comprehensive review stems from the observation that methodological '
        'rigor varies considerably across published work in the field. While some researchers employ '
        'sophisticated statistical frameworks and carefully controlled experimental designs, others rely '
        'on informal evaluation procedures that lack reproducibility. As the discipline matures, '
        'establishing robust methodological standards becomes increasingly critical [i].'
    )

    doc.add_paragraph(
        'Our approach to cataloging research methods follows a taxonomy originally proposed by '
        'Creswell and Creswell (2018), adapted for the specific demands of language technology research. '
        'We distinguish between purely quantitative approaches (corpus frequency analysis, statistical '
        'hypothesis testing, machine learning benchmarks), qualitative approaches (discourse analysis, '
        'interview-based studies, case studies), and mixed-methods designs that integrate both paradigms.'
    )

    doc.add_paragraph(
        'The selection criteria for studies included in this review required peer-reviewed publication '
        'in recognized venues (ACL, EMNLP, NAACL, COLING, Computational Linguistics journal, or '
        'equivalent), a clearly articulated methodology section, and availability of supplementary '
        'materials or code repositories. From an initial pool of 2,847 candidate papers, we retained '
        '412 for detailed analysis after applying exclusion criteria related to scope, completeness, '
        'and methodological transparency.'
    )

    doc.add_paragraph(
        'This document is organized as follows. Section 2 provides a literature review of prior '
        'methodological surveys. Sections 3 and 4 detail quantitative and qualitative methods '
        'respectively. Section 5 discusses mixed-methods designs. Section 6 covers data collection '
        'procedures. Section 7 addresses validity and reliability concerns, and Section 8 presents '
        'our conclusions and recommendations for future research practice.'
    )

    doc.add_page_break()

    # ============================================================
    # CHAPTER 2: LITERATURE REVIEW (pages 4-5)
    # ============================================================
    doc.add_heading('2. Literature Review', level=1)

    doc.add_paragraph(
        'Prior surveys of research methodology in computational linguistics have addressed various '
        'subsets of the methodological landscape. The earliest comprehensive treatment was provided '
        'by Manning and Schutze (1999), who focused primarily on statistical approaches to natural '
        'language processing. Their framework emphasized frequency-based methods, n-gram models, and '
        'hidden Markov models as the dominant analytical tools of the era.'
    )

    doc.add_paragraph(
        'More recently, Jurafsky and Martin (2023) expanded the methodological discussion to '
        'encompass neural approaches, including transformer architectures, attention mechanisms, and '
        'pre-trained language models [ii]. Their treatment, while comprehensive in scope, focused '
        'primarily on algorithmic descriptions rather than experimental methodology per se. The gap '
        'between algorithmic innovation and rigorous experimental evaluation has been noted by several '
        'commentators in the field.'
    )

    doc.add_paragraph(
        'A significant contribution to methodological standards was made by the ACL Rolling Review '
        'initiative, which introduced standardized reviewing criteria that explicitly evaluate '
        'experimental rigor, statistical significance testing, and reproducibility. Analysis of '
        'reviewer scores from 2021-2024 reveals a gradual improvement in methodological quality, '
        'though substantial variation persists across subfields [iii].'
    )

    doc.add_paragraph(
        'The replication crisis in adjacent fields, particularly psychology and biomedical research, '
        'has prompted increased attention to reproducibility in computational linguistics. Dodge et al. '
        '(2019) demonstrated that approximately 63% of published NLP results could not be reproduced '
        'from the information provided in the original papers alone. This finding catalyzed the '
        'establishment of reproducibility tracks at major conferences and the widespread adoption of '
        'code release policies.'
    )

    doc.add_paragraph(
        'Cross-disciplinary methodological borrowing has enriched the toolkit available to '
        'computational linguists. Techniques from psycholinguistics (eye-tracking, reaction time '
        'measurement), sociology (survey design, sampling theory), and information retrieval '
        '(precision-recall frameworks, pooling methodologies) have been adapted for use in language '
        'technology evaluation. This convergence of methods underscores the inherently '
        'interdisciplinary nature of the field.'
    )

    doc.add_page_break()

    # ============================================================
    # CHAPTER 3: QUANTITATIVE METHODS (pages 5-7)
    # ============================================================
    doc.add_heading('3. Quantitative Methods', level=1)

    doc.add_heading('3.1 Corpus Analysis', level=2)

    doc.add_paragraph(
        'Corpus-based methods remain the backbone of empirical computational linguistics. The '
        'construction, annotation, and analysis of text corpora provide the primary evidential basis '
        'for claims about language structure and use. Modern corpus construction practices have evolved '
        'substantially from the early days of the Brown Corpus and the British National Corpus, now '
        'encompassing web-scale data collection, crowdsourced annotation, and semi-automatic quality '
        'control procedures.'
    )

    doc.add_paragraph(
        'The methodological challenges associated with corpus analysis include sampling bias, '
        'annotation reliability, and the ecological validity of derived statistics. Sampling bias '
        'arises when the corpus does not adequately represent the target linguistic population. For '
        'example, web-crawled corpora systematically overrepresent certain registers (news, Wikipedia, '
        'e-commerce) while underrepresenting others (private conversation, technical documentation, '
        'literary fiction) [iv]. Addressing this bias requires careful stratification of data sources '
        'and transparent reporting of corpus composition.'
    )

    doc.add_paragraph(
        'Inter-annotator agreement metrics, particularly Cohen\'s kappa and Krippendorff\'s alpha, '
        'serve as standard measures of annotation reliability. Our review found that 78% of studies '
        'reporting manual annotation included some measure of inter-annotator agreement, though only '
        '42% provided detailed disagreement analysis. Best practices recommend reporting both '
        'agreement statistics and qualitative analysis of disagreement patterns, as numerical '
        'agreement alone may mask systematic biases in annotation guidelines.'
    )

    doc.add_paragraph(
        'The scale of modern corpora introduces both opportunities and challenges. The Common Crawl '
        'dataset, comprising petabytes of web text, enables statistical analyses that were previously '
        'infeasible. However, the sheer volume of data makes manual quality assessment impractical, '
        'necessitating automated filtering and deduplication pipelines. The methodological implications '
        'of these preprocessing steps are often underreported in published work, despite their '
        'potential to significantly influence downstream results.'
    )

    doc.add_heading('3.2 Statistical Modeling', level=2)

    doc.add_paragraph(
        'Statistical modeling in computational linguistics encompasses a broad spectrum of techniques, '
        'from classical hypothesis testing to complex Bayesian hierarchical models. The choice of '
        'statistical framework carries significant implications for the interpretability and '
        'generalizability of research findings.'
    )

    doc.add_paragraph(
        'Frequentist approaches, particularly null hypothesis significance testing (NHST), remain '
        'the most commonly employed statistical framework in the field. Our analysis of 412 papers '
        'found that 67% used NHST as their primary inferential tool. However, the limitations of '
        'NHST, including sensitivity to sample size, the arbitrariness of the p < 0.05 threshold, '
        'and the conflation of statistical and practical significance, have been extensively '
        'documented [v]. Growing awareness of these limitations has prompted increased adoption of '
        'confidence intervals, effect size reporting, and Bayesian alternatives.'
    )

    doc.add_paragraph(
        'Bootstrap resampling methods offer a non-parametric alternative that has gained considerable '
        'traction in NLP evaluation. By repeatedly resampling from the observed data and computing '
        'test statistics on each resample, bootstrap methods provide distribution-free confidence '
        'intervals for performance metrics. This approach is particularly valuable when the sampling '
        'distribution of the test statistic is unknown or when dealing with correlated observations, '
        'as is common in sequence labeling and parsing tasks.'
    )

    doc.add_page_break()

    # ============================================================
    # CHAPTER 4: QUALITATIVE METHODS (pages 7-8)
    # ============================================================
    doc.add_heading('4. Qualitative Methods', level=1)

    doc.add_heading('4.1 Discourse Analysis', level=2)

    doc.add_paragraph(
        'Qualitative methods play an essential though often underappreciated role in computational '
        'linguistics research. Discourse analysis, in particular, provides insights into language use '
        'in context that quantitative methods alone cannot capture. Critical discourse analysis (CDA), '
        'conversation analysis (CA), and multimodal discourse analysis have all found applications in '
        'the study of human-computer interaction, chatbot design, and bias detection in language models.'
    )

    doc.add_paragraph(
        'The integration of qualitative discourse analysis with computational methods has produced '
        'novel hybrid approaches. For instance, topic modeling outputs can be subjected to qualitative '
        'interpretation, combining the scalability of automated analysis with the interpretive depth '
        'of human expertise [vi]. Similarly, error analysis of NLP system outputs frequently employs '
        'qualitative coding schemes to categorize and understand failure modes, providing actionable '
        'insights for system improvement that aggregate metrics cannot offer.'
    )

    doc.add_paragraph(
        'Methodological rigor in qualitative computational linguistics research requires adherence '
        'to established standards of trustworthiness. These include prolonged engagement with the data, '
        'triangulation of findings across multiple data sources or analytical methods, member checking '
        'with research participants (where applicable), and thick description of analytical procedures. '
        'Our review found that qualitative studies in the field exhibited highly variable adherence to '
        'these standards, with only 35% of qualitative papers addressing more than two of the four '
        'trustworthiness criteria.'
    )

    doc.add_heading('4.2 Ethnographic Approaches', level=2)

    doc.add_paragraph(
        'Ethnographic methods, adapted from anthropology and sociology, offer rich descriptions of '
        'how computational linguistic tools are used in practice. Studies employing ethnographic '
        'observation have revealed significant discrepancies between intended and actual use of NLP '
        'applications, highlighting the importance of user-centered evaluation in addition to '
        'benchmark-based assessment.'
    )

    doc.add_paragraph(
        'The ethnographic study of translation technology use by Cadwell et al. (2018) exemplifies '
        'the value of this approach. Through extended observation of professional translators, the '
        'researchers documented how practitioners developed workaround strategies to compensate for '
        'machine translation errors, effectively creating a hybrid human-machine workflow that '
        'differed substantially from the tool designers\' assumptions. Such findings have direct '
        'implications for the design and evaluation of language technology systems.'
    )

    doc.add_paragraph(
        'Participatory design methods represent another qualitative tradition with growing relevance '
        'to computational linguistics. By involving end users in the design process, researchers can '
        'identify requirements and constraints that might otherwise be overlooked. This is particularly '
        'important in applications affecting marginalized communities, where the perspectives of '
        'intended beneficiaries may diverge significantly from those of system developers.'
    )

    doc.add_page_break()

    # ============================================================
    # CHAPTER 5: MIXED METHODS (pages 9-10)
    # ============================================================
    doc.add_heading('5. Mixed Methods Design', level=1)

    doc.add_paragraph(
        'Mixed methods research designs, which systematically integrate quantitative and qualitative '
        'approaches within a single study, offer distinct advantages for computational linguistics '
        'research. By combining the breadth of quantitative analysis with the depth of qualitative '
        'inquiry, mixed methods designs can address research questions that neither approach could '
        'adequately answer alone.'
    )

    doc.add_paragraph(
        'The most common mixed methods design in our review sample was the sequential explanatory '
        'design, in which quantitative analysis is conducted first, followed by qualitative '
        'investigation to explain or elaborate on quantitative findings. For example, a study might '
        'first measure the accuracy of a sentiment analysis system across different text genres '
        '(quantitative phase), then conduct detailed error analysis and user interviews to understand '
        'the sources of genre-specific errors (qualitative phase) [vii].'
    )

    doc.add_paragraph(
        'Convergent parallel designs, in which quantitative and qualitative data are collected '
        'simultaneously and compared during interpretation, were employed in 12% of mixed methods '
        'studies. This approach is particularly useful for validating computational linguistic '
        'analyses against human judgment, as when automated sentiment scores are compared with '
        'qualitative assessments from domain experts.'
    )

    doc.add_paragraph(
        'The challenges of mixed methods research in computational linguistics include the need for '
        'expertise in both quantitative and qualitative paradigms, the complexity of integrating '
        'findings from different analytical traditions, and the additional time and resources required. '
        'Our review identified that mixed methods studies tended to be conducted by larger research '
        'teams (median 5 authors vs. 3 for single-method studies) and were more likely to involve '
        'interdisciplinary collaboration.'
    )

    doc.add_paragraph(
        'The growing emphasis on human-centered AI has created new demand for mixed methods '
        'approaches. Evaluating the societal impact of language technologies requires both '
        'quantitative measurement of system performance and qualitative understanding of user '
        'experience, social dynamics, and ethical implications. We anticipate that mixed methods '
        'designs will become increasingly prevalent as the field grapples with the broader '
        'consequences of deploying language technologies at scale.'
    )

    doc.add_paragraph(
        'Integration of quantitative and qualitative findings presents unique challenges in the '
        'computational linguistics context. Unlike fields where mixed methods integration occurs at '
        'a conceptual level, computational linguistics often requires technical integration at the '
        'data level. For instance, qualitative coding of error types may need to be mapped to '
        'quantitative feature spaces for downstream analysis, requiring careful operationalization '
        'of qualitative categories.'
    )

    doc.add_page_break()

    # ============================================================
    # CHAPTER 6: DATA COLLECTION (pages 10-11)
    # ============================================================
    doc.add_heading('6. Data Collection Procedures', level=1)

    doc.add_paragraph(
        'Data collection in computational linguistics spans a wide range of methods, from automated '
        'web scraping to carefully controlled experimental elicitation. The choice of data collection '
        'method has profound implications for the validity of subsequent analyses and the '
        'generalizability of research findings.'
    )

    doc.add_paragraph(
        'Web scraping and automated data collection remain the most common approaches for building '
        'large-scale corpora. Legal and ethical considerations, including copyright, terms of service '
        'compliance, and privacy protection, have received increasing attention following the '
        'enactment of GDPR and similar data protection regulations. Best practices now include '
        'obtaining institutional review board approval, implementing data anonymization pipelines, '
        'and maintaining transparent documentation of data provenance.'
    )

    doc.add_paragraph(
        'Crowdsourced data collection through platforms such as Amazon Mechanical Turk, Prolific, '
        'and Appen has become a standard method for obtaining human judgments, annotations, and '
        'linguistic data. However, concerns about annotation quality, worker demographics, and '
        'fair compensation have led to the development of more rigorous crowdsourcing protocols. '
        'Qualification tests, attention checks, and iterative guideline refinement are now '
        'considered essential components of crowdsourcing methodology.'
    )

    doc.add_paragraph(
        'Experimental elicitation methods, borrowed from psycholinguistics, provide controlled '
        'environments for studying specific linguistic phenomena. These include acceptability '
        'judgment tasks, sentence completion tasks, picture description tasks, and reading time '
        'measurements. The advantage of experimental methods lies in their ability to isolate '
        'specific variables, though at the cost of ecological validity [viii].'
    )

    doc.add_paragraph(
        'The increasing use of synthetic data generated by large language models presents novel '
        'methodological considerations. While synthetic data can address data scarcity problems, '
        'particularly for low-resource languages and sensitive domains, it may also introduce '
        'systematic biases reflecting the training data and objectives of the generating model. '
        'Rigorous validation against naturally occurring data is essential when synthetic data '
        'is used for training or evaluation purposes.'
    )

    doc.add_paragraph(
        'Longitudinal data collection, tracking language use over extended time periods, has become '
        'more feasible with the growth of social media and digital communication platforms. Studies '
        'of language change, semantic shift, and evolving discourse patterns benefit from temporally '
        'structured corpora. Methodological challenges include ensuring consistent sampling across '
        'time periods, accounting for platform-specific effects, and managing the computational '
        'demands of processing large temporal datasets.'
    )

    doc.add_page_break()

    # ============================================================
    # CHAPTER 7: VALIDITY AND RELIABILITY (pages 11-12)
    # ============================================================
    doc.add_heading('7. Validity and Reliability', level=1)

    doc.add_paragraph(
        'Questions of validity and reliability are central to the credibility of computational '
        'linguistics research. Internal validity concerns whether observed effects can be attributed '
        'to the hypothesized causes rather than confounding factors. External validity addresses '
        'the generalizability of findings beyond the specific experimental conditions. Construct '
        'validity examines whether the measured quantities actually capture the intended theoretical '
        'constructs.'
    )

    doc.add_paragraph(
        'In the context of NLP system evaluation, internal validity threats include data '
        'contamination (test data appearing in training sets), hyperparameter tuning on test data, '
        'and cherry-picking favorable evaluation metrics. Our review found that 23% of papers '
        'did not clearly describe their train-test split methodology, and 41% did not report '
        'results across multiple random seeds or data splits.'
    )

    doc.add_paragraph(
        'External validity concerns are particularly acute in NLP, where performance on benchmark '
        'datasets frequently fails to predict real-world performance. The discrepancy between '
        'benchmark and deployment performance has been attributed to distribution shift, annotation '
        'artifacts, and the limited diversity of standard evaluation datasets. Addressing these '
        'concerns requires evaluation on multiple datasets, cross-domain testing, and user studies '
        'in realistic application contexts.'
    )

    doc.add_paragraph(
        'Reliability in computational linguistics research is assessed through reproducibility '
        'studies and replication attempts. The growing practice of sharing code, data, and '
        'computational environments through platforms such as GitHub, HuggingFace, and Docker '
        'has improved the potential for reproduction, though actual reproduction rates remain low. '
        'Systematic reviews of reproducibility suggest that environmental factors (hardware, library '
        'versions, random seeds) account for a substantial proportion of reproduction failures.'
    )

    doc.add_paragraph(
        'Statistical reliability of evaluation metrics merits particular attention. Small test '
        'sets, imbalanced class distributions, and high-variance metrics can produce unreliable '
        'estimates of system performance. Confidence intervals, significance tests, and multi-split '
        'evaluation protocols help quantify the uncertainty associated with reported results, though '
        'their adoption is not yet universal in the field.'
    )

    doc.add_paragraph(
        'The concept of fairness validity has emerged as an important consideration, examining '
        'whether NLP systems perform equitably across demographic groups. Disparate impact analysis, '
        'calibration testing, and counterfactual evaluation methods have been developed to assess '
        'fairness, complementing traditional validity frameworks with equity considerations.'
    )

    doc.add_page_break()

    # ============================================================
    # CHAPTER 8: CONCLUSIONS (page 12)
    # ============================================================
    doc.add_heading('8. Conclusions', level=1)

    doc.add_paragraph(
        'This survey has examined the principal research methods employed in contemporary '
        'computational linguistics, spanning quantitative, qualitative, and mixed methods approaches. '
        'Our analysis of 412 peer-reviewed studies reveals a field in methodological transition, '
        'with increasing awareness of the limitations of traditional evaluation paradigms and growing '
        'adoption of more rigorous experimental practices.'
    )

    doc.add_paragraph(
        'Several key recommendations emerge from our analysis. First, researchers should routinely '
        'report confidence intervals and effect sizes alongside significance tests, providing a more '
        'complete picture of their findings. Second, the adoption of pre-registration protocols, '
        'already standard in clinical research, would help address concerns about selective reporting '
        'and p-hacking. Third, mixed methods designs should be encouraged when research questions '
        'involve both system performance and human factors.'
    )

    doc.add_paragraph(
        'The field would also benefit from greater methodological diversity. While benchmark-based '
        'evaluation has served computational linguistics well, it should be complemented by '
        'qualitative methods that provide deeper insight into system behavior, user experience, and '
        'societal impact. The development of standardized qualitative evaluation protocols, analogous '
        'to the quantitative benchmarks that have driven progress in the field, represents an '
        'important direction for future methodological work.'
    )

    doc.add_paragraph(
        'Finally, we emphasize the importance of methodological transparency. Detailed reporting of '
        'experimental procedures, data collection methods, statistical analyses, and computational '
        'environments is essential for the cumulative progress of the field. The tools and platforms '
        'for achieving this transparency are now readily available; what remains is to establish the '
        'cultural norms that make transparency the default rather than the exception in computational '
        'linguistics research.'
    )

    doc.add_paragraph(
        'Looking ahead, the rapid pace of technological change in the field presents both '
        'opportunities and challenges for research methodology. The emergence of foundation models, '
        'the increasing scale of experiments, and the growing societal relevance of NLP applications '
        'all call for continued methodological innovation and reflection. We hope that this survey '
        'contributes to that ongoing conversation and supports the development of more rigorous, '
        'transparent, and inclusive research practices in computational linguistics.'
    )

    # Save the document
    doc.save(OUTPUT)

    # Also create a copy with the display name for the task
    import shutil
    shutil.copy(OUTPUT, DISPLAY_NAME)

    print(f'Initial file created: {OUTPUT}')
    print(f'Display copy created: {DISPLAY_NAME}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DISPLAY_NAME}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
