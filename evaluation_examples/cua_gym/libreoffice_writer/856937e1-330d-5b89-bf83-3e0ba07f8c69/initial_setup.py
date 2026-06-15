"""
Initial Setup: Thesis document with Heading 1/2/3 structure, no TOC
Task ID: writer_acad_020
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_020'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title page ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    run = title_para.add_run("Adaptive Machine Learning Approaches for\nClimate Change Impact Assessment")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_before = Pt(36)
    r = subtitle.add_run("A Thesis Submitted in Partial Fulfillment\n"
                         "of the Requirements for the Degree of\n"
                         "Doctor of Philosophy in Computer Science")
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_before = Pt(48)
    r = author.add_run("Elena Vasquez Martinez\nDepartment of Computer Science\nStanford University\nMarch 2025")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    # Page break after title page
    doc.add_page_break()

    # =====================================================================
    # Chapter 1
    # =====================================================================
    doc.add_heading("Introduction", level=1)

    doc.add_heading("Background and Motivation", level=2)
    doc.add_paragraph(
        "Climate change represents one of the most pressing challenges facing humanity in the "
        "twenty-first century. Rising global temperatures, shifting precipitation patterns, and "
        "increasing frequency of extreme weather events demand sophisticated analytical tools "
        "capable of processing vast amounts of environmental data. Traditional statistical methods, "
        "while foundational, often struggle to capture the nonlinear dynamics and complex feedback "
        "loops inherent in Earth's climate system."
    )
    doc.add_paragraph(
        "Machine learning, with its ability to identify patterns in high-dimensional data without "
        "explicit programming of physical rules, has emerged as a promising complement to physics-based "
        "climate models. Recent advances in deep learning architectures, transfer learning, and "
        "ensemble methods have opened new avenues for climate prediction, impact assessment, and "
        "adaptation planning at both regional and global scales."
    )

    doc.add_heading("Research Objectives", level=2)
    doc.add_paragraph(
        "This thesis pursues three primary research objectives. First, we develop an adaptive "
        "neural network framework that dynamically adjusts its architecture based on the spatial "
        "and temporal resolution of input climate data. Second, we design a multi-task learning "
        "pipeline that simultaneously predicts temperature anomalies, precipitation changes, and "
        "sea-level rise for a given geographic region. Third, we propose a novel uncertainty "
        "quantification method that provides calibrated confidence intervals for long-range "
        "climate projections."
    )

    doc.add_heading("Scope and Limitations", level=3)
    doc.add_paragraph(
        "The scope of this work is limited to surface-level climate variables over land masses "
        "in the Northern Hemisphere. Oceanic circulation models and atmospheric chemistry "
        "interactions are treated as external boundary conditions rather than endogenous variables. "
        "While we acknowledge that a fully coupled Earth system model would provide more complete "
        "projections, the computational cost of such integration remains prohibitive for the "
        "iterative hyperparameter searches required by our adaptive framework."
    )

    doc.add_heading("Contributions", level=3)
    doc.add_paragraph(
        "The principal contributions of this thesis are: (1) the AdaptiveClimateNet architecture "
        "that achieves state-of-the-art performance on the CMIP6 benchmark suite; (2) a "
        "calibrated uncertainty framework validated against 40 years of historical reanalysis "
        "data; and (3) an open-source software toolkit, PyClimateML, enabling reproducible "
        "climate impact assessment by the broader research community."
    )

    doc.add_heading("Thesis Organization", level=2)
    doc.add_paragraph(
        "The remainder of this thesis is organized as follows. Chapter 2 reviews the relevant "
        "literature on climate modeling and machine learning. Chapter 3 describes the data sources "
        "and preprocessing pipeline. Chapter 4 presents the AdaptiveClimateNet architecture. "
        "Chapter 5 details the multi-task learning approach. Chapter 6 introduces the uncertainty "
        "quantification framework. Chapter 7 reports experimental results, and Chapter 8 concludes "
        "with a discussion of implications and future research directions."
    )

    doc.add_page_break()

    # =====================================================================
    # Chapter 2
    # =====================================================================
    doc.add_heading("Literature Review", level=1)

    doc.add_heading("Traditional Climate Modeling", level=2)
    doc.add_paragraph(
        "General Circulation Models (GCMs) have served as the backbone of climate science for "
        "over five decades. These physics-based models discretize the atmosphere and ocean into "
        "three-dimensional grids and solve the Navier-Stokes equations subject to radiative "
        "forcing scenarios. The Coupled Model Intercomparison Project (CMIP) has standardized "
        "GCM outputs, enabling systematic comparison across modeling centers worldwide."
    )

    doc.add_heading("Statistical Downscaling Methods", level=3)
    doc.add_paragraph(
        "Statistical downscaling bridges the gap between coarse GCM output (typically 100-250 km "
        "resolution) and the fine-grained projections needed for local impact assessment. Methods "
        "range from simple delta mapping and quantile-quantile correction to more sophisticated "
        "approaches such as stochastic weather generators and analog methods. While computationally "
        "efficient, these techniques assume stationarity of statistical relationships, an assumption "
        "increasingly questioned under nonstationary climate change."
    )

    doc.add_heading("Bias Correction Techniques", level=3)
    doc.add_paragraph(
        "Systematic biases in GCM outputs have motivated the development of bias correction "
        "methods. Quantile mapping, a widely adopted approach, adjusts the cumulative distribution "
        "function of model output to match observations. Recent innovations include "
        "distribution-aware corrections that preserve physical constraints such as non-negative "
        "precipitation and bounded relative humidity."
    )

    doc.add_heading("Machine Learning in Climate Science", level=2)
    doc.add_paragraph(
        "The application of machine learning to climate science has accelerated dramatically since "
        "2015. Convolutional neural networks have been applied to satellite imagery classification, "
        "recurrent neural networks to time-series forecasting of temperature and precipitation, and "
        "generative adversarial networks to super-resolution of climate fields. Graph neural "
        "networks, which naturally accommodate the irregular geometry of observational networks, "
        "represent the latest frontier in data-driven climate modeling."
    )

    doc.add_heading("Deep Learning for Spatial Prediction", level=3)
    doc.add_paragraph(
        "U-Net architectures, originally developed for biomedical image segmentation, have been "
        "adapted for spatial climate prediction with remarkable success. The encoder-decoder "
        "structure with skip connections preserves fine-scale features while capturing large-scale "
        "patterns. Vision Transformers have recently shown competitive performance, leveraging "
        "self-attention mechanisms to model long-range spatial dependencies without the locality "
        "constraints of convolutional kernels."
    )

    doc.add_heading("Uncertainty Quantification in ML Models", level=3)
    doc.add_paragraph(
        "Bayesian neural networks, Monte Carlo dropout, and deep ensembles represent three "
        "complementary strategies for quantifying epistemic and aleatoric uncertainty in ML "
        "predictions. Conformal prediction, a distribution-free framework, has gained traction "
        "for providing finite-sample coverage guarantees without parametric assumptions about "
        "the error distribution."
    )

    doc.add_page_break()

    # =====================================================================
    # Chapter 3
    # =====================================================================
    doc.add_heading("Data and Preprocessing", level=1)

    doc.add_heading("Data Sources", level=2)
    doc.add_paragraph(
        "This study draws on three primary data sources: (1) ERA5 reanalysis data from the "
        "European Centre for Medium-Range Weather Forecasts, spanning 1979 to 2023 at 0.25-degree "
        "spatial resolution; (2) CMIP6 multi-model ensemble projections under SSP2-4.5 and "
        "SSP5-8.5 scenarios; and (3) high-resolution topographic and land-use datasets from "
        "the Copernicus Land Monitoring Service."
    )

    doc.add_heading("Variable Selection", level=3)
    doc.add_paragraph(
        "We select 14 predictor variables spanning thermodynamic, dynamic, and surface "
        "categories. Thermodynamic predictors include 2-meter temperature, specific humidity at "
        "850 hPa, and mean sea-level pressure. Dynamic predictors include 500 hPa geopotential "
        "height, zonal and meridional wind components at 700 hPa, and vertical velocity at "
        "500 hPa. Surface predictors include soil moisture, snow depth, vegetation index, "
        "albedo, elevation, slope, and land-use classification."
    )

    doc.add_heading("Quality Control and Gap Filling", level=3)
    doc.add_paragraph(
        "Raw station observations undergo a multi-stage quality control pipeline. Gross error "
        "checks remove physically implausible values (e.g., temperatures below -90 degrees "
        "Celsius or above 60 degrees Celsius). Spatial consistency checks flag stations whose "
        "anomalies deviate by more than four standard deviations from neighboring stations within "
        "a 200 km radius. Temporal continuity checks identify abrupt jumps exceeding three times "
        "the interquartile range of the station's historical record."
    )

    doc.add_heading("Feature Engineering", level=2)
    doc.add_paragraph(
        "Beyond the raw predictor variables, we engineer 23 derived features to capture known "
        "physical relationships. These include temperature lapse rates computed from multi-level "
        "pressure data, moisture flux convergence, potential vorticity at the tropopause level, "
        "and lagged autocorrelation features at 1-day, 7-day, and 30-day intervals. Principal "
        "component analysis reduces the combined feature space from 37 dimensions to 18 "
        "orthogonal components explaining 95 percent of the total variance."
    )

    doc.add_heading("Train-Validation-Test Split", level=3)
    doc.add_paragraph(
        "To respect the temporal structure of climate data and avoid information leakage, we "
        "employ a chronological split: 1979-2010 for training (32 years), 2011-2017 for "
        "validation (7 years), and 2018-2023 for testing (6 years). Spatial cross-validation "
        "is conducted by holding out entire climate zones, ensuring that the model is evaluated "
        "on its ability to generalize across geographically distinct regions."
    )

    doc.add_page_break()

    # =====================================================================
    # Chapter 4
    # =====================================================================
    doc.add_heading("AdaptiveClimateNet Architecture", level=1)

    doc.add_heading("Design Principles", level=2)
    doc.add_paragraph(
        "AdaptiveClimateNet is built on three design principles: resolution adaptivity, physical "
        "consistency, and computational scalability. Resolution adaptivity allows the network to "
        "process inputs at varying spatial and temporal granularities without retraining. Physical "
        "consistency is enforced through soft constraints that penalize violations of conservation "
        "laws. Computational scalability is achieved via a hierarchical attention mechanism that "
        "reduces the quadratic complexity of standard self-attention to linearithmic complexity."
    )

    doc.add_heading("Encoder Module", level=3)
    doc.add_paragraph(
        "The encoder consists of a stack of multi-scale convolutional blocks, each containing "
        "parallel branches with kernel sizes of 3x3, 5x5, and 7x7. Feature maps from different "
        "scales are concatenated and passed through a 1x1 convolution for channel reduction. "
        "Residual connections and layer normalization stabilize training across the 12-layer "
        "encoder depth."
    )

    doc.add_heading("Adaptive Resolution Layer", level=3)
    doc.add_paragraph(
        "The adaptive resolution layer is the key innovation of our architecture. It employs "
        "learnable interpolation kernels that are conditioned on the input resolution metadata. "
        "When the model receives 0.25-degree ERA5 data, the interpolation kernels produce a "
        "fine-grained feature representation. When processing 1-degree CMIP6 data, the same "
        "kernels smoothly adjust to produce a coarser but physically consistent representation."
    )

    doc.add_heading("Decoder and Output Heads", level=2)
    doc.add_paragraph(
        "The decoder mirrors the encoder with transposed convolutions and skip connections. "
        "Three task-specific output heads produce predictions for temperature anomalies (linear "
        "activation), precipitation changes (softplus activation to ensure non-negativity), and "
        "sea-level rise (linear activation with physically bounded range). Each output head "
        "includes a calibration layer that maps raw predictions to calibrated probability "
        "distributions."
    )

    doc.add_heading("Loss Function Design", level=3)
    doc.add_paragraph(
        "The composite loss function combines mean squared error for point predictions, "
        "continuous ranked probability score for probabilistic calibration, and a physics-informed "
        "penalty term for energy balance violations. The three components are weighted by "
        "task-specific coefficients optimized via Bayesian hyperparameter search on the "
        "validation set."
    )

    doc.add_page_break()

    # =====================================================================
    # Chapter 5
    # =====================================================================
    doc.add_heading("Experimental Results and Discussion", level=1)

    doc.add_heading("Benchmark Comparison", level=2)
    doc.add_paragraph(
        "We compare AdaptiveClimateNet against five baseline methods: linear regression with "
        "principal components, random forests, gradient-boosted trees (XGBoost), a standard "
        "U-Net, and a Vision Transformer. Table 5.1 summarizes performance metrics across all "
        "test regions and seasons. AdaptiveClimateNet achieves the lowest root mean squared "
        "error for temperature prediction (0.73 degrees Celsius versus 0.89 for U-Net and 1.12 "
        "for XGBoost) and the highest Brier skill score for extreme event detection (0.82 versus "
        "0.71 for the Vision Transformer)."
    )

    doc.add_heading("Regional Performance Analysis", level=3)
    doc.add_paragraph(
        "Performance varies substantially across climate zones. The model excels in mid-latitude "
        "continental regions where large-scale circulation patterns dominate local variability. "
        "In tropical regions, where convective processes operate at sub-grid scales, errors "
        "increase by approximately 40 percent. Arctic regions present the greatest challenge due "
        "to sparse observational coverage and rapid nonstationary changes in sea ice extent."
    )

    doc.add_heading("Uncertainty Calibration Results", level=3)
    doc.add_paragraph(
        "The calibration plots in Figure 5.3 demonstrate that our conformal prediction intervals "
        "achieve the target coverage rates across all time horizons. For 1-year projections, the "
        "90 percent prediction intervals contain 91.2 percent of observed values. For 10-year "
        "projections, coverage is 88.7 percent, a slight undercoverage attributable to structural "
        "model uncertainty not captured by the conformal framework."
    )

    doc.add_heading("Ablation Studies", level=2)
    doc.add_paragraph(
        "Ablation experiments reveal the contribution of each architectural component. Removing "
        "the adaptive resolution layer increases RMSE by 18 percent on multi-resolution inputs "
        "while having negligible effect on single-resolution benchmarks. Disabling the physics-"
        "informed penalty term improves training speed by 22 percent but degrades physical "
        "consistency metrics by 31 percent, confirming the value of incorporating domain knowledge."
    )

    doc.add_page_break()

    # =====================================================================
    # Chapter 6
    # =====================================================================
    doc.add_heading("Conclusion and Future Work", level=1)

    doc.add_heading("Summary of Findings", level=2)
    doc.add_paragraph(
        "This thesis has presented AdaptiveClimateNet, a resolution-adaptive deep learning "
        "framework for climate change impact assessment. Our experiments demonstrate that the "
        "proposed architecture achieves state-of-the-art performance across multiple climate "
        "prediction tasks while providing calibrated uncertainty estimates. The open-source "
        "PyClimateML toolkit enables the broader research community to reproduce and extend "
        "these results."
    )

    doc.add_heading("Limitations", level=2)
    doc.add_paragraph(
        "Several limitations merit acknowledgment. The restriction to Northern Hemisphere land "
        "masses excludes Southern Hemisphere and oceanic dynamics that influence global climate "
        "teleconnections. The reliance on ERA5 reanalysis as ground truth introduces biases "
        "inherent in the reanalysis system itself. Furthermore, our uncertainty framework "
        "addresses epistemic uncertainty from finite training data but does not account for "
        "structural uncertainty arising from the choice of emission scenario."
    )

    doc.add_heading("Future Research Directions", level=3)
    doc.add_paragraph(
        "Future work will extend the framework to coupled land-ocean-atmosphere modeling, "
        "incorporate paleoclimate proxy data for training on longer time horizons, and explore "
        "foundation model architectures pretrained on the full CMIP6 archive. Integration with "
        "decision-support systems for climate adaptation planning represents a particularly "
        "promising application domain."
    )

    doc.add_heading("Broader Impact", level=3)
    doc.add_paragraph(
        "As climate change accelerates, the need for accurate, accessible, and interpretable "
        "projection tools will only grow. By democratizing access to state-of-the-art climate "
        "prediction through open-source software and comprehensive documentation, this work "
        "aims to empower researchers, policymakers, and communities to make informed decisions "
        "in the face of an uncertain climate future."
    )

    # Set default font for body paragraphs
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
