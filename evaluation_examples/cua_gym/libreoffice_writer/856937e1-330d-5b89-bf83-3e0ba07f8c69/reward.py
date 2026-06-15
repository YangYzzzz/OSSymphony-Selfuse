"""
Reward Script: Generate a Table of Contents for a thesis with 3 heading levels,
bold Heading 1 entries, and dot leaders with right-aligned page numbers.
Task ID: writer_acad_020
Domain: libreoffice_writer
Scoring:
  Component 1: TOC title exists (0.15)
  Component 2: Three heading levels in TOC (0.20)
  Component 3: Heading 1 entries are bold (0.25)
  Component 4: Dot leaders with right-aligned page numbers (0.25)
  Component 5: Correct number of TOC entries (0.15)
"""

import os
import re
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_020'


def persist_app_state(domain):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify TOC task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    except ImportError as e:
        print(f"CRITICAL: Missing library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # First, collect all headings from the document body (excluding TOC entries)
    # We need these to verify TOC completeness later
    body_headings = {"h1": [], "h2": [], "h3": []}
    for para in doc.paragraphs:
        if para.style.name == "Heading 1":
            body_headings["h1"].append(para.text.strip())
        elif para.style.name == "Heading 2":
            body_headings["h2"].append(para.text.strip())
        elif para.style.name == "Heading 3":
            body_headings["h3"].append(para.text.strip())

    # Identify TOC region: look for a TOC title paragraph followed by TOC entries
    toc_title_idx = None
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.strip().lower()
        if "table of contents" in text_lower or "contents" == text_lower:
            toc_title_idx = i
            break

    # Also check for structured TOC via field codes in XML
    if toc_title_idx is None:
        for i, para in enumerate(doc.paragraphs):
            xml = para._element.xml
            if "TOC" in xml and ("fldChar" in xml or "instrText" in xml):
                toc_title_idx = max(0, i - 1)
                break

    # Component 1: TOC title exists (0.15 points)
    try:
        if toc_title_idx is not None:
            print(f"PASS: Component 1 -- TOC title found at paragraph {toc_title_idx} (0.15 pts)")
            total_score += 0.15
        else:
            print("FAIL: Component 1 -- No 'Table of Contents' title found in document")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    if toc_title_idx is None:
        # No TOC at all - nothing else to verify
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Collect TOC entries: paragraphs after the TOC title that have tab+number pattern
    # TOC entries typically have format "Heading text\tPageNumber"
    toc_entries = []
    for i in range(toc_title_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()

        # Stop at first heading or empty paragraph after entries have been collected
        if para.style.name in ("Heading 1", "Heading 2", "Heading 3"):
            break

        # TOC entry: has a tab character followed by a number
        if "\t" in text and re.search(r'\d+\s*$', text):
            left_indent = para.paragraph_format.left_indent
            indent_val = left_indent if left_indent is not None else 0

            # Determine level by indentation
            # indent=0 -> level 1, indent~228600 -> level 2, indent~457200 -> level 3
            if indent_val <= 50000:
                level = 1
            elif indent_val <= 340000:
                level = 2
            else:
                level = 3

            # Check if entry text (before tab) is bold
            entry_text_part = text.split("\t")[0].strip()
            is_bold = False
            for run in para.runs:
                if run.text.strip() and run.text.strip() in entry_text_part:
                    if run.font.bold is True:
                        is_bold = True
                        break

            # Check tab stops for dot leaders
            has_dot_leader = False
            has_right_align = False
            for ts in para.paragraph_format.tab_stops:
                if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
                    continue
                if ts.alignment == WD_TAB_ALIGNMENT.RIGHT:
                    has_right_align = True
                if ts.leader == WD_TAB_LEADER.DOTS:
                    has_dot_leader = True

            toc_entries.append({
                "text": entry_text_part,
                "level": level,
                "is_bold": is_bold,
                "has_dot_leader": has_dot_leader,
                "has_right_align": has_right_align,
                "indent": indent_val,
            })
        elif text == "" and len(toc_entries) > 0:
            # Empty paragraph after TOC entries - end of TOC
            break

    print(f"\nTOC entries found: {len(toc_entries)}")
    level_counts = {1: 0, 2: 0, 3: 0}
    for e in toc_entries:
        level_counts[e["level"]] = level_counts.get(e["level"], 0) + 1
    print(f"  Level 1 entries: {level_counts.get(1, 0)}")
    print(f"  Level 2 entries: {level_counts.get(2, 0)}")
    print(f"  Level 3 entries: {level_counts.get(3, 0)}")

    # Component 2: Three heading levels in TOC (0.20 points)
    try:
        levels_present = sum(1 for lv in [1, 2, 3] if level_counts.get(lv, 0) > 0)
        if levels_present == 3:
            print(f"PASS: Component 2 -- All 3 heading levels present in TOC (0.20 pts)")
            total_score += 0.20
        elif levels_present == 2:
            print(f"PARTIAL: Component 2 -- Only {levels_present}/3 heading levels in TOC (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 -- Only {levels_present}/3 heading levels in TOC")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Heading 1 (level 1) entries are bold (0.25 points)
    try:
        level1_entries = [e for e in toc_entries if e["level"] == 1]
        if len(level1_entries) == 0:
            print("FAIL: Component 3 -- No level 1 entries found to check bold")
        else:
            bold_count = sum(1 for e in level1_entries if e["is_bold"])
            non_level1_bold = sum(1 for e in toc_entries if e["level"] != 1 and e["is_bold"])

            if bold_count == len(level1_entries) and non_level1_bold == 0:
                print(f"PASS: Component 3 -- All {bold_count}/{len(level1_entries)} level 1 entries are bold, "
                      f"no non-level-1 entries are bold (0.25 pts)")
                total_score += 0.25
            elif bold_count == len(level1_entries):
                # All level 1 bold, but some others too - still mostly correct
                print(f"PARTIAL: Component 3 -- All level 1 entries bold, but {non_level1_bold} non-level-1 "
                      f"entries also bold (0.15 pts)")
                total_score += 0.15
            elif bold_count > 0:
                ratio = bold_count / len(level1_entries)
                pts = round(0.25 * ratio, 2)
                print(f"PARTIAL: Component 3 -- {bold_count}/{len(level1_entries)} level 1 entries are bold ({pts} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 3 -- No level 1 entries are bold (0/{len(level1_entries)})")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Dot leaders with right-aligned page numbers (0.25 points)
    try:
        if len(toc_entries) == 0:
            print("FAIL: Component 4 -- No TOC entries to check tab stops")
        else:
            dot_leader_count = sum(1 for e in toc_entries if e["has_dot_leader"])
            right_align_count = sum(1 for e in toc_entries if e["has_right_align"])

            dot_ratio = dot_leader_count / len(toc_entries)
            right_ratio = right_align_count / len(toc_entries)

            if dot_ratio >= 0.9 and right_ratio >= 0.9:
                print(f"PASS: Component 4 -- {dot_leader_count}/{len(toc_entries)} entries have dot leaders, "
                      f"{right_align_count}/{len(toc_entries)} have right-aligned tabs (0.25 pts)")
                total_score += 0.25
            elif dot_ratio >= 0.5 or right_ratio >= 0.5:
                pts = round(0.25 * max(dot_ratio, right_ratio), 2)
                print(f"PARTIAL: Component 4 -- dot leaders: {dot_leader_count}/{len(toc_entries)}, "
                      f"right-aligned: {right_align_count}/{len(toc_entries)} ({pts} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 4 -- dot leaders: {dot_leader_count}/{len(toc_entries)}, "
                      f"right-aligned: {right_align_count}/{len(toc_entries)}")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    # Component 5: Correct number of TOC entries (0.15 points)
    # The TOC should have entries for all 3 levels of headings in the document
    try:
        expected_total = len(body_headings["h1"]) + len(body_headings["h2"]) + len(body_headings["h3"])
        actual_total = len(toc_entries)

        if actual_total == 0:
            print(f"FAIL: Component 5 -- No TOC entries (expected ~{expected_total})")
        elif abs(actual_total - expected_total) <= 2:
            # Allow small tolerance for edge cases
            print(f"PASS: Component 5 -- TOC has {actual_total} entries (expected ~{expected_total}) (0.15 pts)")
            total_score += 0.15
        elif actual_total >= expected_total * 0.7:
            print(f"PARTIAL: Component 5 -- TOC has {actual_total} entries (expected ~{expected_total}) (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 5 -- TOC has {actual_total} entries (expected ~{expected_total})")
    except Exception as e:
        print(f"ERROR: Component 5 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
