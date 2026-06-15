"""
FINAL REWARD SCRIPT - SUCCESS
Task: Paragraph 4 is one long string—“Name;Role;Email;Phone;Office” and a few more fields—all separated by semicolons. I’m trying to break that line into a real table in LibreOffice Writer, with ‘;’ as the exact delimiter so each bit lands in its own column. How do I pull that off?
Generated: 2025-09-10 13:39:50
Status: success
Model: azure-o3
Total Steps: 2
"""

import os
import re
from docx import Document

def verify_break_line_into_table(file_path: str) -> float:
    """Reward script for verifying that a long semicolon-separated line was
    correctly converted into a real table in LibreOffice Writer / DOCX.

    Scoring (progressive, max 1.0):
        • 0.5 – A table exists that contains at least 5 columns (one for each
          expected field).
        • 0.5 – First-row texts exactly match the expected headers
          ["Name", "Role", "Email", "Phone", "Office"].

    No points are awarded for natural conditions such as file existence or
    successful loading. Those are treated as prerequisites only.
    """

    max_score = 1.0
    score = 0.0

    expected_headers = ["Name", "Role", "Email", "Phone", "Office"]

    # ---------------------------- Prerequisite Checks ----------------------------
    if not os.path.exists(file_path):
        print("✗ File not found. Task cannot be verified.")
        print("REWARD: 0.0")
        return 0.0  # No progress possible

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Unable to load document: {e}")
        print("REWARD: 0.0")
        return 0.0

    print(f"✓ Document loaded -> paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}")

    # ---------------------------- Requirement 1 ----------------------------------
    # A table with at least the required number of columns must exist.
    candidate_table = None
    for t_idx, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if rows else 0
        print(f"  – Found table {t_idx + 1}: size {rows}×{cols}")
        if cols >= len(expected_headers):
            candidate_table = table
            break  # Use the first qualifying table

    if candidate_table is None:
        print("✗ No table with sufficient columns found (need ≥ 5). 0 points awarded for table existence.")
    else:
        score += 0.5
        print("✓ Table with ≥ 5 columns detected (0.5 points)")

        # ------------------------ Requirement 2 ----------------------------------
        # First row (header) must match expected headers exactly (case-insensitive).
        header_cells = [cell.text.strip() for cell in candidate_table.rows[0].cells]
        print("  Header row texts:", header_cells)

        headers_match = True
        for idx, expected in enumerate(expected_headers):
            if idx >= len(header_cells):
                headers_match = False
                break
            if header_cells[idx].strip().lower() != expected.lower():
                headers_match = False
                break

        if headers_match:
            score += 0.5
            print("✓ Header row matches expected labels (0.5 points)")
        else:
            print("✗ Header row does not match expected labels. 0 points for header accuracy.")

    # ---------------------------- Informational Check ----------------------------
    # OPTIONAL: Warn if the raw semicolon string still exists as a paragraph
    semicolon_pattern = re.compile(r"[^\n]*;[^\n]*;")
    leftover_found = False
    for para in doc.paragraphs:
        if semicolon_pattern.search(para.text):
            leftover_found = True
            print(f"⚠︎ Raw semicolon line still present: '{para.text[:80]}…'")
            break
    if not leftover_found:
        print("✓ No raw semicolon-separated line detected in paragraphs (informational)")

    # ---------------------------- Final Score ------------------------------------
    final_score = min(score, max_score)
    print(f"REWARD: {final_score}")
    return final_score


# ------------------------------------------------------------------------------
# Execute verification using the known path provided in the task context.
# ------------------------------------------------------------------------------
DOCX_PATH = "/home/user/paragraph_4_is_one_long_stringnameroleemailphoneoffice_and_a_few_more_fieldsall_separated_by_semicol.docx"
verify_break_line_into_table(DOCX_PATH)
