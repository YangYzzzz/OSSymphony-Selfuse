"""
Initial Setup: Create a thesis document with five key terms appearing in multiple locations.
Task ID: writer_acad_054
Domain: libreoffice_writer
No index entries are marked. No alphabetical index is present.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_054'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ========== TITLE PAGE ==========
    # Add blank paragraphs for vertical centering
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_heading('Adaptive Computational Methods for Predictive Analytics in Healthcare', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy')
    run.font.size = Pt(14)

    doc.add_paragraph()

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('Elena Vasquez')
    run.font.size = Pt(14)
    run.bold = True

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = dept.add_run('Department of Computer Science\nStanford University\nMarch 2026')
    run.font.size = Pt(12)

    doc.add_page_break()

    # ========== ABSTRACT ==========
    doc.add_heading('Abstract', level=1)

    doc.add_paragraph(
        'This thesis investigates the application of machine learning techniques to '
        'predictive healthcare analytics. We develop novel approaches that leverage '
        'deep learning architectures to improve diagnostic accuracy across multiple '
        'clinical domains. Our work demonstrates that neural network models, when '
        'properly trained on large-scale medical datasets, can achieve performance '
        'comparable to experienced clinicians in classification tasks involving '
        'radiological imaging and pathology slides.'
    )

    doc.add_paragraph(
        'We further explore how regression models can be used to predict patient '
        'outcomes such as length of hospital stay, readmission risk, and treatment '
        'response. Through extensive experimentation, we show that combining machine '
        'learning with domain expertise leads to more robust and interpretable models. '
        'Our deep learning framework incorporates attention mechanisms that provide '
        'clinicians with visual explanations of model predictions.'
    )

    doc.add_page_break()

    # ========== CHAPTER 1: INTRODUCTION ==========
    doc.add_heading('Chapter 1: Introduction', level=1)

    doc.add_heading('1.1 Background and Motivation', level=2)

    doc.add_paragraph(
        'The intersection of artificial intelligence and healthcare has seen remarkable '
        'progress in recent years. Machine learning, a subfield of artificial intelligence, '
        'has emerged as a powerful tool for extracting patterns from complex medical data. '
        'The ability of these algorithms to learn from experience without being explicitly '
        'programmed makes them particularly suited to the heterogeneous nature of clinical data.'
    )

    doc.add_paragraph(
        'Among the various machine learning paradigms, deep learning has garnered '
        'particular attention due to its success in processing unstructured data such as '
        'medical images, clinical notes, and genomic sequences. A neural network, the '
        'fundamental building block of deep learning systems, consists of interconnected '
        'layers of artificial neurons that progressively extract higher-level features '
        'from raw input data.'
    )

    doc.add_heading('1.2 Problem Statement', level=2)

    doc.add_paragraph(
        'Despite the promise of these technologies, several challenges remain. '
        'Classification of rare diseases suffers from severe class imbalance, making '
        'standard classification approaches ineffective. Similarly, regression models '
        'for continuous outcome prediction often fail to capture the nonlinear '
        'relationships inherent in biological systems. This thesis addresses these '
        'challenges through novel architectural designs and training strategies.'
    )

    doc.add_heading('1.3 Research Questions', level=2)

    doc.add_paragraph(
        'This thesis investigates the following questions:'
    )
    doc.add_paragraph(
        'How can deep learning architectures be adapted to handle class-imbalanced '
        'classification problems in medical imaging?',
        style='List Number'
    )
    doc.add_paragraph(
        'What neural network designs are most effective for multi-modal clinical '
        'data fusion?',
        style='List Number'
    )
    doc.add_paragraph(
        'Can regression techniques based on machine learning outperform traditional '
        'statistical models in predicting patient outcomes?',
        style='List Number'
    )

    doc.add_page_break()

    # ========== CHAPTER 2: LITERATURE REVIEW ==========
    doc.add_heading('Chapter 2: Literature Review', level=1)

    doc.add_heading('2.1 Machine Learning in Healthcare', level=2)

    doc.add_paragraph(
        'The application of machine learning to healthcare has a rich history dating '
        'back to the 1970s. Early expert systems attempted to codify medical knowledge '
        'into rule-based frameworks. However, these systems were brittle and required '
        'extensive manual knowledge engineering. The advent of statistical machine learning '
        'in the 1990s brought more flexible approaches, including decision trees, support '
        'vector machines, and ensemble methods.'
    )

    doc.add_paragraph(
        'Modern machine learning approaches have shifted toward data-driven methods '
        'that automatically discover relevant features. Random forests and gradient '
        'boosting machines have proven effective for tabular clinical data, particularly '
        'in classification of disease subtypes and risk stratification. These methods '
        'handle missing data gracefully and provide feature importance rankings that '
        'aid clinical interpretation.'
    )

    doc.add_heading('2.2 Deep Learning Revolution', level=2)

    doc.add_paragraph(
        'The deep learning revolution, catalyzed by the availability of large datasets '
        'and GPU computing, has transformed the landscape of medical AI. Convolutional '
        'neural network architectures have achieved superhuman performance in specific '
        'imaging tasks, including diabetic retinopathy screening and skin lesion '
        'classification. These neural network models learn hierarchical representations '
        'that capture both local texture patterns and global structural features.'
    )

    doc.add_paragraph(
        'Recurrent neural network variants, particularly Long Short-Term Memory (LSTM) '
        'networks, have shown promise in temporal clinical data analysis. These models '
        'can process variable-length sequences of patient encounters, laboratory results, '
        'and medication records to predict future health events. The deep learning '
        'paradigm has also enabled end-to-end learning, eliminating the need for manual '
        'feature engineering that was previously required.'
    )

    doc.add_heading('2.3 Regression and Prediction Models', level=2)

    doc.add_paragraph(
        'Regression analysis forms the backbone of predictive modeling in clinical '
        'research. Traditional linear regression and logistic regression have long been '
        'the workhorses of epidemiological studies. However, the assumption of linearity '
        'limits their ability to capture complex dose-response relationships and '
        'interaction effects.'
    )

    doc.add_paragraph(
        'Neural network-based regression approaches offer greater flexibility by '
        'learning nonlinear mappings between input features and continuous outcomes. '
        'Recent work has shown that deep learning regression models can predict '
        'biomarker levels, disease progression trajectories, and survival times with '
        'improved accuracy compared to classical regression methods. The integration '
        'of machine learning with causal inference frameworks has further enhanced '
        'the interpretability of these predictions.'
    )

    doc.add_page_break()

    # ========== CHAPTER 3: METHODOLOGY ==========
    doc.add_heading('Chapter 3: Methodology', level=1)

    doc.add_heading('3.1 Data Collection and Preprocessing', level=2)

    doc.add_paragraph(
        'Our study utilized three publicly available medical datasets: the MIMIC-III '
        'critical care database, the CheXpert chest radiograph dataset, and the UK '
        'Biobank cohort. Data preprocessing involved standardization of numerical '
        'features, encoding of categorical variables, and imputation of missing values '
        'using multiple imputation by chained equations (MICE).'
    )

    doc.add_heading('3.2 Model Architecture', level=2)

    doc.add_paragraph(
        'We propose a multi-branch neural network architecture that processes different '
        'data modalities through specialized pathways before fusion. The imaging branch '
        'employs a deep learning backbone based on EfficientNet-B4, pretrained on '
        'ImageNet and fine-tuned on medical images. The tabular branch uses a '
        'transformer-based architecture that handles heterogeneous clinical features.'
    )

    doc.add_paragraph(
        'For classification tasks, the fused representation passes through a softmax '
        'layer that produces probability distributions over disease categories. We '
        'address class imbalance through focal loss and oversampling of minority '
        'classes. For regression tasks, the output layer uses a linear activation '
        'function with Huber loss to provide robust estimation of continuous outcomes.'
    )

    doc.add_heading('3.3 Training Protocol', level=2)

    doc.add_paragraph(
        'All neural network models were trained using the AdamW optimizer with a cosine '
        'annealing learning rate schedule. We employed five-fold cross-validation to '
        'estimate generalization performance. Hyperparameter optimization was conducted '
        'using Bayesian optimization over a predefined search space. The deep learning '
        'models were implemented in PyTorch and trained on NVIDIA A100 GPUs.'
    )

    doc.add_paragraph(
        'Transfer learning played a critical role in our approach. Pretrained machine '
        'learning models from the natural image domain were adapted to medical imaging '
        'through progressive unfreezing of layers. This strategy proved particularly '
        'effective for classification of rare conditions where labeled data is scarce. '
        'We also explored self-supervised pretraining as an alternative to ImageNet '
        'initialization, finding it beneficial for regression tasks on chest radiographs.'
    )

    doc.add_page_break()

    # ========== CHAPTER 4: RESULTS ==========
    doc.add_heading('Chapter 4: Results', level=1)

    doc.add_heading('4.1 Classification Performance', level=2)

    doc.add_paragraph(
        'Our multi-modal classification system achieved an area under the ROC curve '
        '(AUC) of 0.943 on the CheXpert test set, surpassing the previous '
        'state-of-the-art by 2.1 percentage points. The deep learning model '
        'demonstrated particularly strong performance on rare pathologies, where '
        'the focal loss and oversampling strategy proved critical. Classification '
        'accuracy on the five most common conditions exceeded 90% for all categories.'
    )

    doc.add_paragraph(
        'Comparison with baseline machine learning methods (random forest, gradient '
        'boosting, logistic regression) confirmed the superiority of neural network '
        'approaches for image-based classification. However, for tabular-only '
        'classification tasks, gradient boosting achieved comparable performance, '
        'suggesting that the advantage of deep learning is most pronounced when '
        'unstructured data is involved.'
    )

    doc.add_heading('4.2 Regression Results', level=2)

    doc.add_paragraph(
        'The regression models for length-of-stay prediction achieved a mean absolute '
        'error (MAE) of 1.23 days on the MIMIC-III test set. Our neural network '
        'regression approach outperformed linear regression by 34% and random forest '
        'regression by 18% in terms of MAE. The attention mechanism provided '
        'interpretable feature attributions that aligned with clinical domain knowledge.'
    )

    doc.add_paragraph(
        'For biomarker prediction, the deep learning regression model achieved '
        'correlation coefficients above 0.85 for 12 out of 15 target biomarkers. '
        'The machine learning ensemble approach combining neural network predictions '
        'with gradient boosting regression achieved the best overall performance, '
        'suggesting that model diversity remains important even in the deep learning era.'
    )

    doc.add_page_break()

    # ========== CHAPTER 5: CONCLUSION ==========
    doc.add_heading('Chapter 5: Conclusion', level=1)

    doc.add_paragraph(
        'This thesis has demonstrated the potential of machine learning and deep '
        'learning methods for predictive healthcare analytics. Our contributions '
        'include a novel multi-modal neural network architecture, effective strategies '
        'for handling class-imbalanced classification, and improved regression models '
        'for continuous clinical outcome prediction.'
    )

    doc.add_paragraph(
        'The results confirm that deep learning approaches, when combined with '
        'appropriate training strategies and domain knowledge, can significantly '
        'advance the state of clinical prediction. Our classification models achieve '
        'expert-level performance on radiological diagnosis, while our regression '
        'models provide accurate and interpretable predictions of patient outcomes.'
    )

    doc.add_paragraph(
        'Future work will focus on federated machine learning approaches that enable '
        'multi-institutional model training without sharing sensitive patient data. '
        'We also plan to extend our neural network architecture to incorporate '
        'longitudinal data and explore the use of large language models for clinical '
        'note analysis. The combination of deep learning with causal inference '
        'represents a promising direction for advancing beyond correlation-based '
        'regression toward truly predictive healthcare AI.'
    )

    doc.add_page_break()

    # ========== REFERENCES ==========
    doc.add_heading('References', level=1)

    refs = [
        'Esteva, A., et al. (2017). Dermatologist-level classification of skin cancer with deep neural networks. Nature, 542(7639), 115-118.',
        'Rajpurkar, P., et al. (2017). CheXNet: Radiologist-level pneumonia detection on chest X-rays with deep learning. arXiv preprint arXiv:1711.05225.',
        'Tomaszewski, M. R., & Gillies, R. J. (2021). The biological meaning of radiomic features. Radiology, 298(3), 505-516.',
        'Johnson, A. E., et al. (2016). MIMIC-III: A freely accessible critical care database. Scientific Data, 3, 160035.',
        'LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep learning. Nature, 521(7553), 436-444.',
        'Beam, A. L., & Kohane, I. S. (2018). Big data and machine learning in health care. JAMA, 319(13), 1317-1318.',
        'Obermeyer, Z., & Emanuel, E. J. (2016). Predicting the future: Big data, machine learning, and clinical medicine. New England Journal of Medicine, 375(13), 1216-1219.',
        'Topol, E. J. (2019). High-performance medicine: The convergence of human and artificial intelligence. Nature Medicine, 25(1), 44-56.',
        'Harutyunyan, H., et al. (2019). Multitask learning and benchmarking with clinical time series data. Scientific Data, 6, 96.',
        'Shickel, B., et al. (2018). Deep EHR: A survey of recent advances in deep learning techniques for electronic health record analysis. IEEE JBHI, 22(5), 1589-1604.',
    ]

    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        run = p.add_run(f'[{i}] {ref}')
        run.font.size = Pt(10)

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
