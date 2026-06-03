"""
Reward Script: Create alphabetical subject index with marked index entries
Task ID: writer_acad_054
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.50): XE index entry fields exist for all 5 required terms (0.10 each)
  - Component 2 (0.20): Sufficient number of XE fields (terms appear multiple times in doc)
  - Component 3 (0.15): Alphabetical Index heading present at end of document
  - Component 4 (0.15): INDEX field code present to generate the index listing
"""

import os
import re
from collections import Counter

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_054'
REQUIRED_TERMS = {'machine learning', 'neural network', 'deep learning', 'classification', 'regression'}


def persist_app_state(domain: str):
    """Attempt to save any unsaved state in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body = doc.element.body

    # =========================================================================
    # Component 1: XE index entry fields exist for all 5 required terms (0.50)
    # Each term contributes 0.10 points.
    # This FAILS on initial (0 XE fields) and PASSES on golden (77 XE fields).
    # =========================================================================
    try:
        all_instr = body.findall('.//w:instrText', ns)
        term_counts = Counter()
        for field in all_instr:
            if field.text and 'XE' in field.text:
                # Extract the term inside quotes: XE "term"
                m = re.search(r'XE\s+"([^"]+)"', field.text)
                if m:
                    term_counts[m.group(1).lower()] += 1

        comp1_score = 0.0
        for term in REQUIRED_TERMS:
            if term_counts.get(term, 0) > 0:
                comp1_score += 0.10
                print(f"PASS: Component 1 — XE entries for '{term}' found ({term_counts[term]} occurrences)")
            else:
                print(f"FAIL: Component 1 — No XE entries found for '{term}'")

        if comp1_score > 0:
            total_score += comp1_score
            print(f"PASS: Component 1 — {comp1_score:.2f}/0.50 pts (terms with XE entries)")
        else:
            print(f"FAIL: Component 1 — No XE entries for any required term (0.00/0.50)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # =========================================================================
    # Component 2: Sufficient XE field count (0.20)
    # The terms appear multiple times in the document, so there should be
    # multiple XE entries (at least 10 total across all terms).
    # This FAILS on initial (0 XE fields) and PASSES on golden (77 XE fields).
    # =========================================================================
    try:
        total_xe = sum(term_counts.values())
        if total_xe >= 10:
            print(f"PASS: Component 2 — {total_xe} total XE fields (>= 10 required) (0.20 pts)")
            total_score += 0.20
        elif total_xe >= 5:
            partial = 0.10
            print(f"PARTIAL: Component 2 — {total_xe} total XE fields (5-9 range, partial credit) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {total_xe} total XE fields (need >= 10)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =========================================================================
    # Component 3: Alphabetical Index heading present at end of document (0.15)
    # A heading paragraph containing "Index" should appear near the end.
    # This FAILS on initial (no such heading) and PASSES on golden.
    # =========================================================================
    try:
        index_heading_found = False
        # Check the last 10 paragraphs for an index heading
        last_paras = doc.paragraphs[-10:] if len(doc.paragraphs) >= 10 else doc.paragraphs
        for para in last_paras:
            style_name = para.style.name if para.style else ''
            text_lower = para.text.strip().lower()
            if 'heading' in style_name.lower() and 'index' in text_lower:
                index_heading_found = True
                print(f"PASS: Component 3 — Index heading found: '{para.text.strip()}' (style: {style_name}) (0.15 pts)")
                total_score += 0.15
                break

        if not index_heading_found:
            print(f"FAIL: Component 3 — No heading with 'Index' found in last 10 paragraphs")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # =========================================================================
    # Component 4: INDEX field code present (0.15)
    # The document must contain an INDEX field that generates the actual index.
    # This FAILS on initial (no INDEX field) and PASSES on golden.
    # =========================================================================
    try:
        index_field_found = False
        for field in all_instr:
            if field.text and re.search(r'\bINDEX\b', field.text):
                index_field_found = True
                print(f"PASS: Component 4 — INDEX field found: '{field.text.strip()}' (0.15 pts)")
                total_score += 0.15
                break

        if not index_field_found:
            print(f"FAIL: Component 4 — No INDEX field code found in document")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
