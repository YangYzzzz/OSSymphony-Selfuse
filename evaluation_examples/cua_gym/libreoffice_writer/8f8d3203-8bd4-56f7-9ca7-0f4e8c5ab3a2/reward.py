"""
Reward Script: Apply heading numbering with Roman numerals for H1 and Arabic for H2
Task ID: writer_list_061
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.4): All Heading 1 paragraphs have numbering (numPr) at ilvl=0
  - Component 2 (0.4): All Heading 2 paragraphs have numbering (numPr) at ilvl=1
  - Component 3 (0.2): Numbering uses upperRoman for level 0 and decimal with %1.%2 for level 1
"""

import os

from docx import Document
from docx.oxml.ns import qn
import lxml.etree as etree

WORKDIR = '/home/user'
TASK_ID = 'writer_list_061'


def get_numId_ilvl(paragraph):
    """Return (numId, ilvl) for a paragraph, or (None, None) if no numbering."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        return None, None
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return None, None
    ilvl_el = numPr.find(qn('w:ilvl'))
    numId_el = numPr.find(qn('w:numId'))
    ilvl = int(ilvl_el.get(qn('w:val'))) if ilvl_el is not None else None
    numId = int(numId_el.get(qn('w:val'))) if numId_el is not None else None
    return numId, ilvl


def get_abstractNum_for_numId(doc, numId):
    """Resolve numId to its abstractNum element in the numbering part.
    When multiple w:num entries share the same numId (can happen after heading numbering
    is applied on top of existing list styles), return the LAST one — which is the
    newly added heading numbering definition.
    """
    try:
        num_part = doc.part.numbering_part
        if num_part is None:
            return None
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        root = num_part._element
        # Find ALL <w:num w:numId="X"> elements and take the last one
        matching_num_els = [
            num_el for num_el in root.findall('w:num', ns)
            if num_el.get(qn('w:numId')) == str(numId)
        ]
        if not matching_num_els:
            return None
        # Use the last match (the most recently added — heading numbering)
        num_el = matching_num_els[-1]
        abstract_id_el = num_el.find('w:abstractNumId', ns)
        if abstract_id_el is None:
            return None
        abstract_id = abstract_id_el.get(qn('w:val'))
        # Find the abstractNum with that abstractNumId
        for abstract_el in root.findall('w:abstractNum', ns):
            if abstract_el.get(qn('w:abstractNumId')) == abstract_id:
                return abstract_el
        return None
    except Exception as e:
        print(f"ERROR: get_abstractNum_for_numId: {e}")
        return None


def get_level_info(abstract_el, ilvl):
    """Return (numFmt, lvlText) for a given ilvl in an abstractNum element."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for lvl_el in abstract_el.findall('w:lvl', ns):
        if lvl_el.get(qn('w:ilvl')) == str(ilvl):
            numFmt_el = lvl_el.find('w:numFmt', ns)
            lvlText_el = lvl_el.find('w:lvlText', ns)
            numFmt = numFmt_el.get(qn('w:val')) if numFmt_el is not None else None
            lvlText = lvlText_el.get(qn('w:val')) if lvlText_el is not None else None
            return numFmt, lvlText
    return None, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task is to apply heading numbering:
    - Heading 1 → uppercase Roman numerals (I, II, III) via ilvl=0
    - Heading 2 → Arabic with parent prefix (I.1, I.2) via ilvl=1
    Both levels should share the same numId pointing to a multi-level abstractNum.
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect headings
    h1_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 1']
    h2_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 2']

    print(f"INFO: Found {len(h1_paras)} Heading 1 paragraphs and {len(h2_paras)} Heading 2 paragraphs")

    # Expected heading texts from context
    expected_h1 = ['Executive Summary', 'Research Methodology', 'Findings and Analysis']
    expected_h2 = ['Data Collection Methods', 'Sample Size and Demographics',
                   'Statistical Analysis Approach', 'Quantitative Results', 'Qualitative Observations']

    # Component 1: All Heading 1 paragraphs have numbering (numPr) at ilvl=0 (0.4 points)
    try:
        h1_numbered_count = 0
        h1_ilvl_correct = 0
        for p in h1_paras:
            numId, ilvl = get_numId_ilvl(p)
            if numId is not None and numId != 0:
                h1_numbered_count += 1
                if ilvl == 0:
                    h1_ilvl_correct += 1
                    print(f"PASS: Heading 1 '{p.text[:40]}' has numId={numId}, ilvl={ilvl}")
                else:
                    print(f"FAIL: Heading 1 '{p.text[:40]}' has numId={numId} but ilvl={ilvl} (expected 0)")
            else:
                print(f"FAIL: Heading 1 '{p.text[:40]}' has no numbering (numId={numId})")

        if len(h1_paras) > 0 and h1_ilvl_correct == len(h1_paras):
            print(f"PASS: Component 1 — All {len(h1_paras)} Heading 1 paragraphs have numbering at ilvl=0 (0.4 pts)")
            total_score += 0.4
        elif h1_ilvl_correct > 0:
            partial = round(0.4 * h1_ilvl_correct / len(h1_paras), 2)
            print(f"PARTIAL: Component 1 — {h1_ilvl_correct}/{len(h1_paras)} Heading 1 paragraphs have correct numbering ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No Heading 1 paragraphs have numbering at ilvl=0")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All Heading 2 paragraphs have numbering (numPr) at ilvl=1 (0.4 points)
    try:
        h2_ilvl_correct = 0
        for p in h2_paras:
            numId, ilvl = get_numId_ilvl(p)
            if numId is not None and numId != 0:
                if ilvl == 1:
                    h2_ilvl_correct += 1
                    print(f"PASS: Heading 2 '{p.text[:40]}' has numId={numId}, ilvl={ilvl}")
                else:
                    print(f"FAIL: Heading 2 '{p.text[:40]}' has numId={numId} but ilvl={ilvl} (expected 1)")
            else:
                print(f"FAIL: Heading 2 '{p.text[:40]}' has no numbering (numId={numId})")

        if len(h2_paras) > 0 and h2_ilvl_correct == len(h2_paras):
            print(f"PASS: Component 2 — All {len(h2_paras)} Heading 2 paragraphs have numbering at ilvl=1 (0.4 pts)")
            total_score += 0.4
        elif h2_ilvl_correct > 0:
            partial = round(0.4 * h2_ilvl_correct / len(h2_paras), 2)
            print(f"PARTIAL: Component 2 — {h2_ilvl_correct}/{len(h2_paras)} Heading 2 paragraphs have correct numbering ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No Heading 2 paragraphs have numbering at ilvl=1")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Numbering format — upperRoman for level 0, decimal with %1.%2 for level 1 (0.2 points)
    try:
        # Find the numId used by Heading 1 paragraphs
        h1_numId = None
        for p in h1_paras:
            numId, ilvl = get_numId_ilvl(p)
            if numId is not None and numId != 0:
                h1_numId = numId
                break

        if h1_numId is None:
            print("FAIL: Component 3 — Cannot check format because no Heading 1 has numbering")
        else:
            abstract_el = get_abstractNum_for_numId(doc, h1_numId)
            if abstract_el is None:
                print(f"FAIL: Component 3 — Cannot resolve abstractNum for numId={h1_numId}")
            else:
                # Check level 0: should be upperRoman
                numFmt_0, lvlText_0 = get_level_info(abstract_el, 0)
                # Check level 1: should be decimal with text %1.%2 (or similar with Roman+Arabic)
                numFmt_1, lvlText_1 = get_level_info(abstract_el, 1)

                print(f"INFO: Level 0 — numFmt={numFmt_0!r}, lvlText={lvlText_0!r}")
                print(f"INFO: Level 1 — numFmt={numFmt_1!r}, lvlText={lvlText_1!r}")

                # Level 0 must use upperRoman
                level0_ok = (numFmt_0 == 'upperRoman')
                # Level 1 must use decimal
                level1_fmt_ok = (numFmt_1 == 'decimal')
                # Level 1 text should include the parent Roman numeral reference (%1) and Arabic numeral (%2)
                # Expected patterns: "%1.%2" or "I.%2" — most common is "%1.%2"
                level1_text_ok = (lvlText_1 is not None and '%1' in lvlText_1 and '%2' in lvlText_1)

                if level0_ok and level1_fmt_ok and level1_text_ok:
                    print(f"PASS: Component 3 — Level 0 is upperRoman, Level 1 is decimal with parent reference (0.2 pts)")
                    total_score += 0.2
                elif level0_ok:
                    print(f"PARTIAL: Component 3 — Level 0 is upperRoman but Level 1 format incorrect: numFmt={numFmt_1}, lvlText={lvlText_1} (0.1 pts)")
                    total_score += 0.1
                else:
                    print(f"FAIL: Component 3 — Level 0 format is {numFmt_0!r}, expected 'upperRoman'")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
