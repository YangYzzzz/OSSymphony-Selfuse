"""
Initial Setup: Formal report with Heading 1 and Heading 2 styled paragraphs (no numbering)
Task ID: writer_list_061
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_list_061'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Executive Summary (Heading 1) ---
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report presents a comprehensive analysis of market trends and operational "
        "performance indicators for the fiscal year 2024. The findings reveal significant "
        "growth opportunities in emerging markets, particularly in the Asia-Pacific region. "
        "Key performance metrics indicate a 12.4% increase in overall revenue compared to "
        "the previous year, driven primarily by digital transformation initiatives."
    )
    doc.add_paragraph(
        "The executive team should prioritize investments in technology infrastructure and "
        "talent acquisition to capitalize on the identified opportunities. Risk mitigation "
        "strategies must address supply chain vulnerabilities and regulatory compliance "
        "requirements across multiple jurisdictions."
    )

    # --- Research Methodology (Heading 1) ---
    doc.add_heading("Research Methodology", level=1)
    doc.add_paragraph(
        "The research was conducted over a six-month period from January to June 2024, "
        "employing a mixed-methods approach to ensure comprehensive data coverage. Both "
        "quantitative surveys and qualitative interviews were utilized to capture diverse "
        "perspectives from stakeholders across the organization."
    )

    # --- Data Collection Methods (Heading 2) ---
    doc.add_heading("Data Collection Methods", level=2)
    doc.add_paragraph(
        "Primary data was collected through structured online surveys distributed to "
        "3,847 participants across 12 regional offices. Survey instruments were designed "
        "using validated scales from peer-reviewed literature and adapted to reflect "
        "industry-specific requirements. Response rate achieved was 78.3%, exceeding "
        "the target threshold of 70%."
    )
    doc.add_paragraph(
        "Secondary data sources included annual reports from 45 comparable organizations, "
        "industry benchmarking databases, and publicly available financial disclosures. "
        "Data triangulation methods were applied to validate findings across multiple sources."
    )

    # --- Sample Size and Demographics (Heading 2) ---
    doc.add_heading("Sample Size and Demographics", level=2)
    doc.add_paragraph(
        "The study sample comprised 3,012 valid respondents after data cleaning and "
        "exclusion of incomplete responses. Demographic distribution included 58% male "
        "and 42% female participants, with age groups ranging from 22 to 67 years. "
        "Geographic representation spanned North America (34%), Europe (28%), "
        "Asia-Pacific (26%), and other regions (12%)."
    )
    doc.add_paragraph(
        "Stratified random sampling was employed to ensure proportional representation "
        "from each business unit. Minimum sample sizes per stratum were calculated using "
        "Cochran's formula with a 95% confidence level and 3% margin of error."
    )

    # --- Statistical Analysis Approach (Heading 2) ---
    doc.add_heading("Statistical Analysis Approach", level=2)
    doc.add_paragraph(
        "Quantitative data analysis was performed using SPSS version 28.0 and R version "
        "4.3.1. Descriptive statistics were computed for all variables, followed by "
        "inferential analyses including multiple regression, factor analysis, and "
        "structural equation modeling. All tests were conducted at the 0.05 significance level."
    )
    doc.add_paragraph(
        "Qualitative data from semi-structured interviews was analyzed using thematic "
        "analysis following the Braun and Clarke (2006) framework. NVivo 14 software "
        "facilitated the coding process, with inter-rater reliability assessed through "
        "Cohen's kappa coefficient (κ = 0.84)."
    )

    # --- Findings and Analysis (Heading 1) ---
    doc.add_heading("Findings and Analysis", level=1)
    doc.add_paragraph(
        "The research yielded several significant findings that have important implications "
        "for strategic planning and resource allocation. Both quantitative metrics and "
        "qualitative insights converge on key themes related to operational efficiency, "
        "market positioning, and organizational resilience."
    )

    # --- Quantitative Results (Heading 2) ---
    doc.add_heading("Quantitative Results", level=2)
    doc.add_paragraph(
        "Revenue analysis revealed total earnings of $4.7 billion for fiscal year 2024, "
        "representing a 12.4% year-over-year growth rate. Operating margin improved from "
        "18.2% to 21.7%, attributed to operational efficiency programs implemented in Q2 "
        "and Q3. Customer acquisition cost decreased by 8.3% while customer lifetime value "
        "increased by 15.6%, indicating improved marketing effectiveness."
    )
    doc.add_paragraph(
        "Employee productivity metrics showed a 9.2% improvement in output per full-time "
        "equivalent, correlated strongly (r = 0.73, p < 0.001) with adoption rates of new "
        "digital tools. Net Promoter Score rose from 42 to 58 points, placing the "
        "organization in the top quartile for its industry segment."
    )

    # --- Qualitative Observations (Heading 2) ---
    doc.add_heading("Qualitative Observations", level=2)
    doc.add_paragraph(
        "In-depth interviews with 47 senior managers revealed recurring themes around "
        "change management challenges and cultural adaptation requirements. Participants "
        "consistently emphasized the importance of transparent communication during "
        "organizational transitions, particularly regarding technology adoption roadmaps "
        "and role evolution expectations."
    )
    doc.add_paragraph(
        "Three distinct employee experience archetypes emerged from the qualitative data: "
        "early adopters (23%), pragmatic followers (54%), and resistant traditionalists "
        "(23%). Each group requires differentiated engagement strategies to maximize "
        "participation in digital transformation initiatives and minimize productivity "
        "disruption during transition periods."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
