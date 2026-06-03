"""
Initial Setup: Legal contract document without visual separator between preamble and operative provisions
Task ID: writer_legal_056
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_056'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Date and parties
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Effective Date: March 15, 2025')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacer

    # Opening recital
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    run = intro.add_run(
        'This Professional Services Agreement (the "Agreement") is entered into as of '
        'March 15, 2025 (the "Effective Date"), by and between:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # Party A
    party_a = doc.add_paragraph()
    party_a.paragraph_format.left_indent = Inches(0.5)
    party_a.paragraph_format.space_after = Pt(6)
    run = party_a.add_run(
        'Meridian Technologies, Inc., a Delaware corporation with its principal office '
        'located at 2400 Innovation Drive, Suite 800, San Jose, CA 95134 '
        '(hereinafter referred to as the "Company")'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # and
    and_para = doc.add_paragraph()
    and_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    and_para.paragraph_format.space_after = Pt(6)
    run = and_para.add_run('and')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # Party B
    party_b = doc.add_paragraph()
    party_b.paragraph_format.left_indent = Inches(0.5)
    party_b.paragraph_format.space_after = Pt(12)
    run = party_b.add_run(
        'Apex Consulting Group, LLC, a California limited liability company with its '
        'principal office located at 1750 Market Street, Floor 12, San Francisco, CA 94102 '
        '(hereinafter referred to as the "Consultant")'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # WHEREAS clauses (preamble / recitals)
    whereas_clauses = [
        (
            'WHEREAS, the Company is engaged in the business of developing and marketing '
            'enterprise software solutions for financial institutions and requires specialized '
            'consulting services in connection with its next-generation platform migration;'
        ),
        (
            'WHEREAS, the Consultant possesses extensive expertise in cloud infrastructure '
            'architecture, data migration strategies, and enterprise system integration, and '
            'has previously delivered similar services to Fortune 500 clients;'
        ),
        (
            'WHEREAS, the Company desires to engage the Consultant to provide certain '
            'professional services as described herein, and the Consultant desires to provide '
            'such services, subject to the terms and conditions set forth in this Agreement;'
        ),
        (
            'WHEREAS, both parties acknowledge that the successful completion of the '
            'Platform Migration Project (as defined below) is critical to the Company\'s '
            'strategic objectives for fiscal year 2025-2026 and requires dedicated resources '
            'and coordinated effort;'
        ),
    ]

    for clause_text in whereas_clauses:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.first_line_indent = Inches(0.5)
        run = para.add_run(clause_text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # NOW, THEREFORE — flows directly after WHEREAS with NO visual separator
    now_para = doc.add_paragraph()
    now_para.paragraph_format.space_before = Pt(12)
    now_para.paragraph_format.space_after = Pt(6)
    now_para.paragraph_format.first_line_indent = Inches(0.5)
    run = now_para.add_run(
        'NOW, THEREFORE, in consideration of the mutual covenants, promises, and '
        'agreements contained herein, and for other good and valuable consideration, '
        'the receipt and sufficiency of which are hereby acknowledged, the parties '
        'agree as follows:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # Article 1 - Definitions
    art1 = doc.add_heading('ARTICLE 1 — DEFINITIONS', level=1)
    art1.paragraph_format.space_before = Pt(18)

    definitions = [
        ('"Deliverables"', ' means all work product, reports, documentation, software code, '
         'and other materials created by the Consultant in the course of performing the Services.'),
        ('"Project Plan"', ' means the detailed timeline, milestones, and resource allocation '
         'document attached hereto as Exhibit A.'),
        ('"Services"', ' means the professional consulting services described in Article 2 of '
         'this Agreement and any Statements of Work executed hereunder.'),
        ('"Confidential Information"', ' means any non-public information disclosed by either '
         'party to the other, whether orally, in writing, or in electronic form, that is designated '
         'as confidential or that reasonably should be understood to be confidential.'),
    ]

    for term, definition in definitions:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Inches(0.5)
        run_term = para.add_run(term)
        run_term.bold = True
        run_term.font.size = Pt(11)
        run_term.font.name = 'Times New Roman'
        run_def = para.add_run(definition)
        run_def.font.size = Pt(11)
        run_def.font.name = 'Times New Roman'

    # Article 2 - Scope of Services
    art2 = doc.add_heading('ARTICLE 2 — SCOPE OF SERVICES', level=1)
    art2.paragraph_format.space_before = Pt(18)

    scope_text = (
        'The Consultant shall provide the following professional services to the Company '
        'in accordance with the Project Plan and any applicable Statements of Work:'
    )
    scope_para = doc.add_paragraph()
    scope_para.paragraph_format.space_after = Pt(6)
    run = scope_para.add_run(scope_text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    services = [
        'Assessment and documentation of the Company\'s existing legacy infrastructure, '
        'including hardware inventory, software dependencies, and data flow mappings.',
        'Design and architecture of the target cloud-native platform on Amazon Web Services (AWS), '
        'including network topology, security groups, and auto-scaling configurations.',
        'Development of a phased migration strategy with rollback procedures and '
        'disaster recovery planning for each migration phase.',
        'Knowledge transfer sessions and technical documentation for the Company\'s '
        'internal engineering team upon completion of each project milestone.',
    ]

    for svc in services:
        para = doc.add_paragraph(svc, style='List Number')
        para.paragraph_format.space_after = Pt(4)
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # Article 3 - Compensation
    art3 = doc.add_heading('ARTICLE 3 — COMPENSATION', level=1)
    art3.paragraph_format.space_before = Pt(18)

    comp_para = doc.add_paragraph()
    comp_para.paragraph_format.space_after = Pt(6)
    run = comp_para.add_run(
        'In consideration of the Services to be performed by the Consultant, the Company '
        'shall pay the Consultant a total fee of Four Hundred Seventy-Five Thousand Dollars '
        '($475,000.00), payable in accordance with the following schedule:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # Payment table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Milestone', 'Amount', 'Due Date']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

    payments = [
        ['Project Kickoff', '$95,000.00', 'April 1, 2025'],
        ['Phase 1 Completion', '$142,500.00', 'June 30, 2025'],
        ['Phase 2 Completion', '$142,500.00', 'September 30, 2025'],
        ['Final Acceptance', '$95,000.00', 'December 15, 2025'],
    ]
    for r, row_data in enumerate(payments, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
