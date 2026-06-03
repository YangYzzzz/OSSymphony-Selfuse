"""
Initial Setup: Create a Writer document with a heading and five plain text feature lines.
Task ID: writer_rd_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a heading
    heading = doc.add_heading("Key Features", level=1)

    # Five plain-text feature paragraphs (no list formatting, no bullets)
    features = [
        "Seamless integration with over 200 third-party applications and cloud services, enabling real-time data synchronization across your entire workflow.",
        "Advanced analytics dashboard with customizable widgets that provide actionable insights into team performance, project milestones, and resource allocation.",
        "End-to-end encryption for all stored documents and communications, ensuring enterprise-grade security compliance with SOC 2 and GDPR standards.",
        "Intelligent task prioritization powered by machine learning that adapts to your work patterns and automatically surfaces high-impact items each morning.",
        "Collaborative editing with granular version control, allowing up to 50 simultaneous users to work on the same document without merge conflicts.",
    ]

    for feature_text in features:
        para = doc.add_paragraph(feature_text)
        # Ensure plain Normal style, no list formatting
        para.style = doc.styles['Normal']

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer (GUI-ready state)
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
