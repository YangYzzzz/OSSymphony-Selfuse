"""
Initial Setup: Create an office memo with 10 ALL CAPS words for Find & Replace task
Task ID: writer_frd_019
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_019'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Title --
    title = doc.add_heading('Internal Memo', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -- Metadata block --
    meta = doc.add_paragraph()
    meta.add_run('To: ').bold = True
    meta.add_run('All Department Heads\n')
    run_from = meta.add_run('From: ')
    run_from.bold = True
    meta.add_run('Sarah Mitchell, Chief Operations Officer\n')
    run_date = meta.add_run('Date: ')
    run_date.bold = True
    meta.add_run('March 28, 2026\n')
    run_subj = meta.add_run('Subject: ')
    run_subj.bold = True
    meta.add_run('URGENT \u2014 End-of-Quarter Action Items and COMPLIANCE Review')
    # ALL CAPS words in this paragraph: URGENT (1), COMPLIANCE (2)

    # -- Horizontal line --
    doc.add_paragraph('_' * 60)

    # -- Greeting --
    doc.add_paragraph('Dear Colleagues,')
    doc.add_paragraph()

    # -- Paragraph 1: IMPORTANT, DEADLINE --
    p1 = doc.add_paragraph(
        'As we approach the end of Q1 2026, there are several IMPORTANT '
        'matters that require your immediate attention. The DEADLINE for '
        'submitting all departmental reports is April 4, 2026. Please ensure '
        'your teams are on track to meet this date without exception.'
    )
    # ALL CAPS: IMPORTANT (3), DEADLINE (4)

    # -- Paragraph 2: BUDGET --
    p2 = doc.add_paragraph(
        'The BUDGET allocation for Q2 has been finalized by the finance team. '
        'Each department will receive a detailed breakdown by the end of next week. '
        'Please review the numbers carefully and flag any discrepancies to '
        'David Park in Accounting.'
    )
    # ALL CAPS: BUDGET (5)

    # -- Paragraph 3: MARKETING, STRATEGY --
    p3 = doc.add_paragraph(
        'The MARKETING department has proposed a new outreach campaign targeting '
        'enterprise clients. This aligns with our broader STRATEGY to expand into '
        'the B2B segment. Rachel Torres will present the full plan during the '
        'leadership meeting on April 7.'
    )
    # ALL CAPS: MARKETING (6), STRATEGY (7)

    # -- Paragraph 4: EXECUTIVE, QUARTERLY --
    p4 = doc.add_paragraph(
        'Following the recent audit, the EXECUTIVE leadership team has mandated '
        'a thorough review of all vendor contracts. The QUARTERLY performance '
        'metrics indicate strong growth in three of our five regions, but we need '
        'to address the underperforming areas before the board meeting.'
    )
    # ALL CAPS: EXECUTIVE (8), QUARTERLY (9)

    # -- Paragraph 5: REVENUE --
    p5 = doc.add_paragraph(
        'REVENUE projections for the second half of the year remain optimistic. '
        'However, achieving these targets will require sustained effort and '
        'cross-departmental coordination. Please ensure your teams are aligned '
        'with the goals outlined in the strategic plan distributed last month.'
    )
    # ALL CAPS: REVENUE (10)

    # -- Closing --
    doc.add_paragraph()
    doc.add_paragraph(
        'Please treat these items with the highest priority. If you have '
        'questions, do not hesitate to reach out to my office directly.'
    )

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.add_run('Best regards,\n')
    name_run = closing.add_run('Sarah Mitchell\n')
    name_run.bold = True
    closing.add_run('Chief Operations Officer\n')
    closing.add_run('Meridian Solutions Group')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
