"""
Reward Script: Convert ALL CAPS words to Title Case in memo
Task ID: writer_frd_019
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): No ALL CAPS words (3+ letters) remain
  Component 2 (0.4): All 10 expected Title Case words are present
  Component 3 (0.2): Document text integrity preserved (paragraph count, overall content)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_019'

# The 10 ALL CAPS words from the initial document and their expected Title Case forms
CAPS_TO_TITLE = {
    'URGENT': 'Urgent',
    'COMPLIANCE': 'Compliance',
    'IMPORTANT': 'Important',
    'DEADLINE': 'Deadline',
    'BUDGET': 'Budget',
    'MARKETING': 'Marketing',
    'STRATEGY': 'Strategy',
    'EXECUTIVE': 'Executive',
    'QUARTERLY': 'Quarterly',
    'REVENUE': 'Revenue',
}


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all text from paragraphs
    full_text = ' '.join(para.text for para in doc.paragraphs)

    # Component 1: No ALL CAPS words (3+ letters) remain (0.4 points)
    # This checks the core task requirement: all caps words should be gone.
    # On initial_env, there are 10 caps words -> FAIL
    # On golden_env, there are 0 caps words -> PASS
    try:
        caps_words_found = re.findall(r'\b[A-Z]{3,}\b', full_text)
        if len(caps_words_found) == 0:
            print(f"PASS: Component 1 - No ALL CAPS words remain (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 - Found {len(caps_words_found)} ALL CAPS words: {caps_words_found}")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: All 10 Title Case replacements are present (0.4 points)
    # Each correct replacement earns 0.04 points (10 x 0.04 = 0.4)
    # On initial_env, words are UPPERCASE so Title Case versions won't be found -> 0 pts
    # On golden_env, all 10 should be present -> 0.4 pts
    try:
        title_score = 0.0
        found_count = 0
        for caps_word, title_word in CAPS_TO_TITLE.items():
            # Use word boundary to find exact Title Case word
            pattern = r'\b' + re.escape(title_word) + r'\b'
            if re.search(pattern, full_text):
                found_count += 1
                title_score += 0.04
                print(f"  PASS: '{title_word}' found in document")
            else:
                print(f"  FAIL: '{title_word}' not found in document")

        if found_count > 0:
            print(f"PASS: Component 2 - {found_count}/10 Title Case words found ({title_score:.2f} pts)")
            total_score += title_score
        else:
            print(f"FAIL: Component 2 - No Title Case replacements found")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Document text integrity preserved (0.2 points)
    # The document should still have the same paragraph structure and content
    # (minus the case changes). This verifies the task didn't corrupt the document.
    # On initial_env, paragraphs have CAPS words so lowered comparison would match,
    # but we specifically check that the Title Case words replaced CAPS words.
    # We check: paragraph count is 14 AND key phrases still exist.
    # The key: on initial_env, the title case words DON'T exist, so this component
    # requires BOTH structural integrity AND the presence of title case words.
    try:
        num_paras = len(doc.paragraphs)
        # Check paragraph count (should be 14 based on initial document)
        para_count_ok = (num_paras >= 12 and num_paras <= 16)

        # Check key phrases that should exist in both initial and golden
        # BUT also require at least one title case replacement to be present
        # to ensure this component only scores the golden state
        key_phrases_present = (
            'Sarah Mitchell' in full_text
            and 'Internal Memo' in full_text
            and 'Dear Colleagues' in full_text
        )

        # This sub-check ensures component 3 only passes on golden_env:
        # At least 5 of the title case words must be present
        has_title_case_words = sum(
            1 for tw in CAPS_TO_TITLE.values()
            if re.search(r'\b' + re.escape(tw) + r'\b', full_text)
        ) >= 5

        if para_count_ok and key_phrases_present and has_title_case_words:
            print(f"PASS: Component 3 - Document integrity preserved with title case conversions ({num_paras} paragraphs) (0.2 pts)")
            total_score += 0.2
        else:
            reasons = []
            if not para_count_ok:
                reasons.append(f"paragraph count {num_paras} outside expected range 12-16")
            if not key_phrases_present:
                reasons.append("key phrases missing")
            if not has_title_case_words:
                reasons.append("title case words not present (pre-task state)")
            print(f"FAIL: Component 3 - {'; '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
