"""
Initial Setup: Thesis manuscript without watermark
Task ID: writer_acad_070
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_070'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title Page ---
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_heading("", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("The Impact of Artificial Intelligence on Urban Transportation Networks:\nA Multi-City Comparative Analysis")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.bold = True

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy in Urban Systems Engineering")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph("")
    doc.add_paragraph("")

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run("Elena Vasquez-Moreno")
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph("")

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = dept.add_run("Department of Civil and Environmental Engineering\nStanford University\nMarch 2025")
    run.font.size = Pt(12)

    # Page break after title
    doc.add_page_break()

    # --- Abstract ---
    abstract_heading = doc.add_heading("Abstract", level=1)
    abstract_text = (
        "This dissertation investigates the transformative effects of artificial intelligence (AI) "
        "technologies on urban transportation networks across five major metropolitan areas: "
        "San Francisco, Singapore, Barcelona, Nairobi, and Seoul. Drawing on a mixed-methods "
        "approach that combines large-scale traffic simulation data, ridership analytics from "
        "public transit agencies, and qualitative interviews with 147 urban planners and "
        "transportation engineers, this study examines how AI-driven optimization algorithms "
        "have reshaped commuter behavior, route efficiency, and infrastructure investment "
        "decisions between 2019 and 2024."
    )
    p = doc.add_paragraph(abstract_text)
    p.paragraph_format.line_spacing = 2.0

    abstract_text2 = (
        "Our findings reveal that cities implementing AI-based traffic signal coordination "
        "experienced an average reduction of 23.7% in peak-hour congestion delays, while "
        "AI-optimized public transit scheduling led to a 15.2% increase in ridership across "
        "all studied networks. However, the study also identifies significant disparities "
        "in AI deployment equity, with lower-income neighborhoods receiving substantially "
        "fewer algorithmic optimizations compared to affluent commercial districts. The "
        "research contributes a novel framework, the Urban AI Integration Index (UAII), "
        "for evaluating the holistic impact of AI interventions on transportation equity "
        "and efficiency."
    )
    p2 = doc.add_paragraph(abstract_text2)
    p2.paragraph_format.line_spacing = 2.0

    keywords = doc.add_paragraph()
    keywords.paragraph_format.space_before = Pt(12)
    run = keywords.add_run("Keywords: ")
    run.bold = True
    keywords.add_run(
        "artificial intelligence, urban transportation, traffic optimization, "
        "public transit, smart cities, transportation equity, algorithmic governance"
    )

    doc.add_page_break()

    # --- Chapter 1: Introduction ---
    doc.add_heading("Chapter 1: Introduction", level=1)

    doc.add_heading("1.1 Background and Motivation", level=2)
    p = doc.add_paragraph(
        "The rapid proliferation of artificial intelligence technologies in urban "
        "infrastructure management represents one of the most significant paradigm shifts "
        "in civil engineering since the advent of computerized traffic modeling in the 1970s "
        "(Mitchell & Zhang, 2021). Contemporary urban transportation networks face "
        "unprecedented challenges: growing populations, aging infrastructure, climate change "
        "mandates, and shifting commuter preferences driven by remote work adoption. "
        "Traditional optimization approaches, rooted in static models and periodic manual "
        "adjustments, have proven increasingly inadequate for managing the dynamic complexity "
        "of modern metropolitan mobility systems."
    )
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)

    p = doc.add_paragraph(
        "Recent advances in deep reinforcement learning (DRL), computer vision, and "
        "real-time data processing have enabled a new generation of adaptive transportation "
        "management systems capable of responding to changing conditions within seconds "
        "rather than weeks (Park et al., 2023). These systems leverage vast sensor networks, "
        "GPS data from connected vehicles, and mobile device location data to construct "
        "real-time digital twins of urban traffic patterns, enabling predictive rather than "
        "reactive management strategies."
    )
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_heading("1.2 Research Questions", level=2)
    questions = [
        "How do AI-driven traffic management systems affect peak-hour congestion metrics across cities with varying infrastructure maturity levels?",
        "What is the relationship between AI-optimized public transit scheduling and ridership patterns across different socioeconomic demographics?",
        "To what extent do current AI deployment strategies in transportation perpetuate or mitigate existing urban equity disparities?",
        "How do institutional governance structures mediate the adoption and effectiveness of AI transportation technologies?",
    ]
    for i, q in enumerate(questions, 1):
        rq = doc.add_paragraph(style="List Number")
        rq.add_run(f"RQ{i}: {q}")
        rq.paragraph_format.line_spacing = 2.0

    doc.add_heading("1.3 Scope and Significance", level=2)
    p = doc.add_paragraph(
        "This research encompasses a five-year longitudinal analysis (2019-2024) of AI "
        "integration in transportation systems across five cities selected for their "
        "diversity in economic development, governance structures, and technological "
        "infrastructure. San Francisco represents a mature Western tech hub with extensive "
        "private-sector AI investment; Singapore exemplifies a centralized smart city "
        "governance model; Barcelona illustrates European regulatory frameworks for AI "
        "deployment; Nairobi provides insights into AI adoption in rapidly urbanizing "
        "African contexts; and Seoul demonstrates East Asian approaches to public-private "
        "AI partnerships in transportation."
    )
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # --- Chapter 2: Literature Review ---
    doc.add_heading("Chapter 2: Literature Review", level=1)

    doc.add_heading("2.1 Historical Evolution of Traffic Management Systems", level=2)
    p = doc.add_paragraph(
        "The evolution of traffic management systems can be broadly categorized into four "
        "generations. First-generation systems (1950s-1970s) relied on fixed-time signal "
        "plans based on historical traffic counts (Webster, 1958). Second-generation "
        "systems (1970s-1990s) introduced actuated signals responsive to real-time detector "
        "data, exemplified by the SCOOT and SCATS platforms (Hunt et al., 1982). "
        "Third-generation systems (1990s-2010s) incorporated centralized optimization "
        "algorithms capable of coordinating signals across corridors and networks "
        "(Gartner et al., 1995). The fourth and current generation leverages machine "
        "learning and artificial intelligence for fully adaptive, predictive control "
        "(Li et al., 2022)."
    )
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_heading("2.2 AI Applications in Urban Mobility", level=2)
    p = doc.add_paragraph(
        "The application of artificial intelligence in urban mobility extends beyond "
        "traffic signal optimization to encompass demand prediction, route optimization, "
        "autonomous vehicle coordination, and multimodal journey planning. Reinforcement "
        "learning approaches have demonstrated particular promise in adaptive signal "
        "control, with recent studies reporting 15-30% improvements in average delay "
        "times compared to traditional actuated control methods (Wei et al., 2021). "
        "Natural language processing has been applied to analyze commuter feedback "
        "and social media data to identify service quality issues in real time "
        "(Chen & Bharadwaj, 2023)."
    )
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_heading("2.3 Equity Considerations in Smart Transportation", level=2)
    p = doc.add_paragraph(
        "A growing body of literature raises concerns about the equitable distribution "
        "of benefits from AI-driven transportation improvements. Dillahunt and Veinot "
        "(2018) documented how algorithmic route optimization can inadvertently redirect "
        "traffic through lower-income residential neighborhoods, increasing pollution "
        "exposure while reducing commute times for wealthier commuters. Similarly, "
        "Eubanks (2018) argues that data-driven resource allocation in public services, "
        "including transportation, often reinforces existing patterns of underinvestment "
        "in marginalized communities."
    )
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # --- Chapter 3: Methodology ---
    doc.add_heading("Chapter 3: Methodology", level=1)

    doc.add_heading("3.1 Research Design", level=2)
    p = doc.add_paragraph(
        "This study employs a convergent parallel mixed-methods design (Creswell & "
        "Plano Clark, 2018), integrating quantitative traffic simulation analysis with "
        "qualitative expert interviews. The quantitative strand utilizes microsimulation "
        "models calibrated with real-world sensor data from each study city, while the "
        "qualitative strand draws on semi-structured interviews with transportation "
        "professionals to contextualize numerical findings within institutional and "
        "cultural frameworks."
    )
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_heading("3.2 Data Collection", level=2)

    # Data sources table
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["City", "Data Sources", "Time Period", "Sample Size"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ["San Francisco", "SFMTA sensors, Uber Movement API, BART ridership logs", "2019-2024", "48M trip records"],
        ["Singapore", "LTA DataMall, ERP gantry data, SimMobility outputs", "2019-2024", "62M trip records"],
        ["Barcelona", "TMB open data, Ajuntament sensors, RACC traffic feeds", "2020-2024", "31M trip records"],
        ["Nairobi", "Ma3Route crowdsourced data, NTSA records, Safaricom GPS", "2021-2024", "8.7M trip records"],
        ["Seoul", "TOPIS real-time data, T-money transaction logs, KOTI models", "2019-2024", "89M trip records"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph("")  # spacing after table

    doc.add_heading("3.3 Analytical Framework", level=2)
    p = doc.add_paragraph(
        "The Urban AI Integration Index (UAII) framework developed in this study "
        "evaluates AI transportation interventions across four dimensions: efficiency "
        "gains (measured by delay reduction and throughput improvement), equity impact "
        "(measured by Gini coefficient changes in service quality distribution), "
        "sustainability outcomes (measured by emissions reduction estimates), and "
        "institutional capacity (measured by governance readiness indicators). Each "
        "dimension is scored on a normalized scale from 0 to 100, with composite "
        "scores enabling cross-city comparison."
    )
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # --- Chapter 4: Results ---
    doc.add_heading("Chapter 4: Results", level=1)

    doc.add_heading("4.1 Congestion Reduction Outcomes", level=2)
    p = doc.add_paragraph(
        "Analysis of peak-hour congestion data reveals statistically significant "
        "reductions in average delay times across all five cities following AI system "
        "deployment. San Francisco demonstrated the largest improvement with a 31.4% "
        "reduction in average intersection delay (p < 0.001), followed by Seoul (27.8%), "
        "Singapore (24.1%), Barcelona (19.3%), and Nairobi (15.9%). These results "
        "remained robust after controlling for concurrent infrastructure investments, "
        "population changes, and seasonal variation using a difference-in-differences "
        "estimation strategy."
    )
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Results table
    results_table = doc.add_table(rows=6, cols=5)
    results_table.style = "Table Grid"
    r_headers = ["City", "Delay Reduction (%)", "Ridership Change (%)", "Emissions Impact (%)", "UAII Score"]
    for i, h in enumerate(r_headers):
        cell = results_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    r_data = [
        ["San Francisco", "31.4", "+18.7", "-12.3", "78.4"],
        ["Seoul", "27.8", "+21.2", "-9.8", "82.1"],
        ["Singapore", "24.1", "+16.5", "-14.7", "85.6"],
        ["Barcelona", "19.3", "+11.8", "-8.2", "71.3"],
        ["Nairobi", "15.9", "+8.4", "-4.1", "54.7"],
    ]
    for r, row_data in enumerate(r_data, 1):
        for c, val in enumerate(row_data):
            results_table.cell(r, c).text = val

    doc.add_paragraph("")

    doc.add_heading("4.2 Equity Analysis", level=2)
    p = doc.add_paragraph(
        "The equity analysis reveals a persistent and concerning pattern across all "
        "study cities: AI-driven transportation improvements disproportionately benefit "
        "higher-income areas. In San Francisco, neighborhoods in the top income quartile "
        "received 3.2 times more AI-optimized signal intersections per capita than those "
        "in the bottom quartile. Singapore presented the most equitable distribution, "
        "attributed to its centralized planning model, though disparities persisted between "
        "central business districts and peripheral housing estates."
    )
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # --- Chapter 5: Discussion ---
    doc.add_heading("Chapter 5: Discussion and Conclusions", level=1)

    doc.add_heading("5.1 Implications for Urban Policy", level=2)
    p = doc.add_paragraph(
        "The findings of this study carry significant implications for urban policy "
        "makers navigating the integration of AI technologies into transportation "
        "infrastructure. First, the demonstrated efficiency gains validate continued "
        "investment in AI-driven traffic management, but the equity disparities identified "
        "underscore the need for explicit equity mandates in AI deployment strategies. "
        "We recommend that municipalities adopt the UAII framework as a standard evaluation "
        "tool to ensure balanced outcomes across socioeconomic strata."
    )
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_heading("5.2 Limitations and Future Research", level=2)
    p = doc.add_paragraph(
        "This study acknowledges several limitations that warrant consideration. The "
        "five-city sample, while diverse, cannot capture the full spectrum of urban "
        "contexts globally. The reliance on available open data introduces potential "
        "selection bias, as cities with more transparent data practices may differ "
        "systematically from those with limited data availability. Future research "
        "should expand the geographic scope to include Latin American, South Asian, "
        "and Middle Eastern cities, and should incorporate longitudinal tracking of "
        "equity metrics following AI system deployment."
    )
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)

    # --- References (partial) ---
    doc.add_page_break()
    doc.add_heading("References", level=1)

    refs = [
        "Chen, L., & Bharadwaj, A. (2023). Natural language processing for real-time transit feedback analysis. Transportation Research Part C, 148, 103982.",
        "Creswell, J. W., & Plano Clark, V. L. (2018). Designing and conducting mixed methods research (3rd ed.). SAGE Publications.",
        "Dillahunt, T. R., & Veinot, T. C. (2018). Getting there: Barriers and facilitators to transportation access in underserved communities. ACM Transactions on Computer-Human Interaction, 25(5), 1-39.",
        "Eubanks, V. (2018). Automating inequality: How high-tech tools profile, police, and punish the poor. St. Martin's Press.",
        "Gartner, N. H., Pooran, F. J., & Andrews, C. M. (1995). Implementation of the OPAC adaptive control strategy in a traffic signal network. Proceedings of the IEEE ITSC, 195-200.",
        "Hunt, P. B., Robertson, D. I., Bretherton, R. D., & Royle, M. C. (1982). The SCOOT on-line traffic signal optimisation technique. Traffic Engineering and Control, 23(4), 190-192.",
        "Li, Z., Yu, H., Zhang, G., Dong, S., & Xu, C. (2022). Network-wide traffic signal control optimization using a multi-agent deep reinforcement learning. Transportation Research Part C, 125, 103059.",
        "Mitchell, T., & Zhang, W. (2021). The evolution of intelligent transportation systems: From automation to autonomy. Annual Review of Control, Robotics, and Autonomous Systems, 4, 345-372.",
        "Park, J., Kim, S., & Lee, D. (2023). Deep reinforcement learning for adaptive traffic signal control: A comprehensive review. IEEE Transactions on Intelligent Transportation Systems, 24(3), 2548-2567.",
        "Webster, F. V. (1958). Traffic signal settings (Road Research Technical Paper No. 39). HMSO, London.",
        "Wei, H., Zheng, G., Gayah, V., & Li, Z. (2021). Recent advances in reinforcement learning for traffic signal control: A survey. ACM Computing Surveys, 54(7), 1-36.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
