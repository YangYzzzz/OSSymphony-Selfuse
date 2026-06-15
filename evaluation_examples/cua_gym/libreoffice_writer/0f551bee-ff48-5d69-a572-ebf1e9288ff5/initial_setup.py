"""
Initial Setup: Enable AutoCorrect numbered list option
Task ID: writer_frd_049
Domain: libreoffice_writer

Initial state: LibreOffice Writer is open with a document. The AutoCorrect
option for 'Apply numbering' (ByInput) is DISABLED (the default).
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_049'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
REGISTRY = f'{WORKDIR}/.config/libreoffice/4/user/registrymodifications.xcu'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def kill_libreoffice():
    """Kill any running LibreOffice processes."""
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    time.sleep(2)


def ensure_apply_numbering_disabled():
    """
    Ensure the ApplyNumbering ByInput option is explicitly set to false
    in registrymodifications.xcu. The default is false, but we set it
    explicitly to be safe.
    """
    import xml.etree.ElementTree as ET

    tree = ET.parse(REGISTRY)
    root = tree.getroot()
    ns = {
        'oor': 'http://openoffice.org/2001/registry',
        'xs': 'http://www.w3.org/2001/XMLSchema',
        'xsi': 'http://www.w3.org/2001/XMLSchema-instance',
    }

    path = "/org.openoffice.Office.Writer/AutoFunction/Format/ByInput/ApplyNumbering"

    # Check if an item with this path already exists
    found = False
    for item in root.findall('item', ns):
        item_path = item.get('{http://openoffice.org/2001/registry}path', '')
        if item_path == path:
            # Update existing entry
            prop = item.find('{http://openoffice.org/2001/registry}prop', ns)
            if prop is None:
                prop = item.find('prop', ns)
            if prop is not None:
                val = prop.find('value')
                if val is None:
                    val = prop.find('{http://openoffice.org/2001/registry}value', ns)
                if val is not None:
                    val.text = 'false'
                    found = True
            break

    if not found:
        # Add new item
        item_tag = '{http://openoffice.org/2001/registry}item'
        # Use the namespace-less approach matching existing file format
        new_item_str = (
            f'<item oor:path="{path}">'
            f'<prop oor:name="Enable" oor:op="fuse">'
            f'<value>false</value>'
            f'</prop></item>'
        )
        # Parse and append
        # Actually, let's just do string manipulation since ElementTree has namespace issues
        pass

    # Simpler approach: use string manipulation on the XML file
    with open(REGISTRY, 'r') as f:
        content = f.read()

    entry = f'<item oor:path="{path}"><prop oor:name="Enable" oor:op="fuse"><value>false</value></prop></item>'

    if path in content:
        # Already has an entry for this path - replace it
        import re
        pattern = rf'<item oor:path="{re.escape(path)}">[^<]*<prop[^>]*>[^<]*<value>[^<]*</value>[^<]*</prop>[^<]*</item>'
        content = re.sub(pattern, entry, content)
    else:
        # Add before closing tag
        content = content.replace('</oor:items>', entry + '\n</oor:items>')

    with open(REGISTRY, 'w') as f:
        f.write(content)

    print(f'ApplyNumbering/Enable set to false in registrymodifications.xcu')


def create_initial_document():
    """Create a simple Writer document with realistic content."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Title
    heading = doc.add_heading('Quarterly Planning Notes', level=1)

    # Some realistic body paragraphs
    doc.add_paragraph(
        'The following items were discussed during the Q2 2025 planning session '
        'held on March 18, 2025. All department leads were present and contributed '
        'to the agenda.'
    )

    doc.add_paragraph(
        'Key priorities for the upcoming quarter include expanding the customer '
        'onboarding workflow, finalizing the vendor contract renewals, and '
        'completing the infrastructure migration to the new cloud platform.'
    )

    doc.add_heading('Action Items', level=2)

    doc.add_paragraph(
        'Below are the action items that need to be tracked. Each team lead '
        'is responsible for updating progress on a weekly basis.'
    )

    doc.add_paragraph(
        'Marketing team will prepare the campaign brief for the product launch '
        'by April 5. The design assets should be reviewed by Sarah Chen before '
        'distribution to external partners.'
    )

    doc.add_paragraph(
        'Engineering is expected to deliver the API integration prototype by '
        'mid-April. Marcus Johnson will coordinate with the QA team to ensure '
        'test coverage meets the 85% threshold.'
    )

    doc.add_paragraph(
        'Finance department needs to reconcile the Q1 expense reports and '
        'submit the updated budget forecast to the CFO by March 28.'
    )

    doc.add_heading('Notes', level=2)

    doc.add_paragraph(
        'The next planning session is scheduled for June 15, 2025. Please '
        'prepare your department updates and any resource requests in advance.'
    )

    doc.save(OUTPUT)
    print(f'Initial document created: {OUTPUT}')


def main():
    # Kill LibreOffice if running so we can modify config
    kill_libreoffice()

    # Ensure the Apply Numbering setting is disabled
    ensure_apply_numbering_disabled()

    # Create the initial document
    create_initial_document()

    # Open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
