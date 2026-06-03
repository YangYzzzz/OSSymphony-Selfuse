"""
Initial Setup: Employee handbook with subsection titles in Default Paragraph Style
Task ID: writer_hr_018
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_018'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Main title - use Heading 1 for the main section
    doc.add_heading('Employee Benefits Handbook', level=1)

    intro = doc.add_paragraph(
        'Welcome to Meridian Technologies. This handbook outlines the comprehensive '
        'benefits package available to all full-time employees. Please review each section '
        'carefully and direct any questions to the Human Resources department.'
    )

    doc.add_paragraph(
        'Meridian Technologies is committed to supporting the well-being of our employees '
        'and their families. Our benefits program is designed to provide financial security, '
        'health coverage, and work-life balance resources that meet your needs at every '
        'stage of your career.'
    )

    # --- Subsection 1: Eligibility (Normal/Default style) ---
    p1 = doc.add_paragraph('Eligibility')
    # Ensure it is Normal style, NOT Heading 2
    p1.style = doc.styles['Normal']

    doc.add_paragraph(
        'All regular full-time employees who work a minimum of 30 hours per week are '
        'eligible for the complete benefits package. Eligibility begins on the first day '
        'of the month following 60 calendar days of continuous employment.'
    )
    doc.add_paragraph(
        'Part-time employees working between 20 and 29 hours per week are eligible for '
        'a limited benefits package that includes dental and vision coverage. Temporary '
        'and contract employees are not eligible for company-sponsored benefits.'
    )
    doc.add_paragraph(
        'Dependents eligible for coverage include your legal spouse or registered domestic '
        'partner, and children up to age 26. Documentation of dependent status may be '
        'required during the annual enrollment verification process.'
    )

    # --- Subsection 2: Coverage Period (Normal/Default style) ---
    p2 = doc.add_paragraph('Coverage Period')
    p2.style = doc.styles['Normal']

    doc.add_paragraph(
        'The standard benefits plan year runs from January 1 through December 31. Open '
        'enrollment occurs annually during the month of November, with changes taking '
        'effect on January 1 of the following year.'
    )
    doc.add_paragraph(
        'Qualifying life events — such as marriage, birth or adoption of a child, '
        'divorce, or loss of other coverage — allow mid-year enrollment changes. '
        'Employees must submit qualifying event documentation within 30 days of the event.'
    )
    doc.add_paragraph(
        'Upon termination of employment, COBRA continuation coverage is available for '
        'up to 18 months. The HR department will provide COBRA election forms within '
        '14 days of the qualifying event.'
    )

    # --- Subsection 3: Claims Process (Normal/Default style) ---
    p3 = doc.add_paragraph('Claims Process')
    p3.style = doc.styles['Normal']

    doc.add_paragraph(
        'Medical claims are processed automatically when you visit an in-network provider '
        'and present your insurance card. For out-of-network services, you must submit a '
        'claim form along with an itemized bill to BlueCross BlueShield within 90 days '
        'of the date of service.'
    )
    doc.add_paragraph(
        'Prescription drug claims are handled through the CVS Caremark pharmacy benefit '
        'manager. Generic prescriptions require a $10 copay, preferred brand-name drugs '
        'require a $30 copay, and non-preferred brand-name drugs require a $50 copay.'
    )
    doc.add_paragraph(
        'Dental claims for preventive services (cleanings, exams, X-rays) are covered '
        'at 100% with no deductible. Basic restorative services are covered at 80% after '
        'the $50 individual deductible, and major services at 50% after deductible.'
    )

    # --- Subsection 4: Appeals (Normal/Default style) ---
    p4 = doc.add_paragraph('Appeals')
    p4.style = doc.styles['Normal']

    doc.add_paragraph(
        'If a claim is denied, you have the right to file an appeal. The first level of '
        'appeal is an internal review by the insurance carrier, which must be initiated '
        'within 180 days of receiving the denial notice.'
    )
    doc.add_paragraph(
        'To file an internal appeal, submit a written request to the carrier along with '
        'any supporting documentation from your healthcare provider. The carrier must '
        'respond to urgent care appeals within 72 hours and standard appeals within 30 days.'
    )
    doc.add_paragraph(
        'If the internal appeal is denied, you may request an external review by an '
        'independent third party. External review decisions are binding on the insurance '
        'carrier. The HR Benefits team can assist you with the appeals process at any stage.'
    )

    # --- Subsection 5: Contact Information (Normal/Default style) ---
    p5 = doc.add_paragraph('Contact Information')
    p5.style = doc.styles['Normal']

    doc.add_paragraph(
        'For general benefits questions, contact the HR Benefits Team at '
        'benefits@meridiantech.com or call extension 4200. Office hours are Monday '
        'through Friday, 8:00 AM to 5:00 PM Eastern Time.'
    )
    doc.add_paragraph(
        'For medical claims inquiries, contact BlueCross BlueShield directly at '
        '1-800-555-0142. Your group number is MT-20250301. For pharmacy benefits, '
        'contact CVS Caremark at 1-800-555-0198.'
    )
    doc.add_paragraph(
        'For dental and vision claims, contact Delta Dental at 1-800-555-0167 and '
        'VSP Vision Care at 1-800-555-0173 respectively. Always have your employee ID '
        'and group number available when calling.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
