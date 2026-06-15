"""
Reward Script: Insert horizontal line separator after memo header (Subject line)
Task ID: writer_biz_007
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Subject paragraph has a bottom border
  Component 2 (0.3): Border style is a visible line (not 'none'/'nil')
  Component 3 (0.2): Body text below the separator is preserved
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_007'


def persist_app_state(domain):
    """Save any unsaved LibreOffice changes before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def find_subject_paragraph(doc):
    """Find the paragraph that starts with 'Subject:' (the last header line)."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.lower().startswith("subject:") or text.lower().startswith("subject\t"):
            return i, para
    return None, None


def has_bottom_border(para):
    """Check if a paragraph has a bottom border element."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return False, None
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        return False, None
    bottom = pBdr.find(qn('w:bottom'))
    if bottom is None:
        return False, None
    return True, bottom


def get_border_style(border_element):
    """Get the border style value (e.g., 'single', 'thick', 'double')."""
    if border_element is None:
        return None
    val = border_element.get(qn('w:val'))
    return val


def has_any_horizontal_separator(doc, subject_idx):
    """
    Check for any kind of horizontal separator after the subject line.
    This could be:
    1. Bottom border on the Subject paragraph
    2. A separate paragraph with borders (top or bottom) right after Subject
    3. A horizontal rule element
    Returns: (found, description)
    """
    # Check 1: Bottom border on Subject paragraph itself
    if subject_idx is not None:
        found, el = has_bottom_border(doc.paragraphs[subject_idx])
        if found:
            return True, "bottom border on Subject paragraph"

    # Check 2: Check the paragraph immediately after Subject for top border or full border
    if subject_idx is not None and subject_idx + 1 < len(doc.paragraphs):
        next_para = doc.paragraphs[subject_idx + 1]
        pPr = next_para._element.find(qn('w:pPr'))
        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                top = pBdr.find(qn('w:top'))
                bottom = pBdr.find(qn('w:bottom'))
                if top is not None or bottom is not None:
                    return True, "border on paragraph after Subject"

    # Check 3: Check if a new paragraph was inserted between Subject and body
    # with border properties (horizontal rule style)
    if subject_idx is not None:
        for idx in range(subject_idx, min(subject_idx + 3, len(doc.paragraphs))):
            para = doc.paragraphs[idx]
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    for side in ['top', 'bottom']:
                        el = pBdr.find(qn(f'w:{side}'))
                        if el is not None:
                            style = el.get(qn('w:val'))
                            if style and style not in ('none', 'nil'):
                                return True, f"{side} border on P{idx}"

    return False, "no separator found"


def _find_visible_border_style(doc, subject_idx):
    """Return the border style string if a visible border is found near Subject, else None."""
    # Check Subject paragraph bottom border
    found_bottom, bottom_el = has_bottom_border(doc.paragraphs[subject_idx])
    if found_bottom:
        style = get_border_style(bottom_el)
        if style and style not in ('none', 'nil'):
            return style
    # Check neighboring paragraphs
    for check_idx in range(subject_idx, min(subject_idx + 3, len(doc.paragraphs))):
        para = doc.paragraphs[check_idx]
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                for side in ['top', 'bottom']:
                    el = pBdr.find(qn(f'w:{side}'))
                    if el is not None:
                        style = get_border_style(el)
                        if style and style not in ('none', 'nil'):
                            return style
    return None


def _body_text_exists_after(doc, subject_idx, expected_start):
    """Return True if body text starting with expected_start exists after subject_idx."""
    for idx in range(subject_idx + 1, len(doc.paragraphs)):
        if expected_start.lower() in doc.paragraphs[idx].text.lower():
            return idx
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate the Subject paragraph
    subject_idx, subject_para = find_subject_paragraph(doc)
    if subject_idx is None:
        print("FAIL: Could not find 'Subject:' paragraph in the document")
        print("REWARD: 0.0")
        return 0.0
    print(f"INFO: Found Subject paragraph at index {subject_idx}: {subject_para.text!r}")

    # Component 1: A horizontal separator exists near the Subject line (0.5 points)
    # This FAILS on initial (no borders) and PASSES on golden (bottom border on Subject)
    try:
        separator_found, separator_desc = has_any_horizontal_separator(doc, subject_idx)
        if separator_found:
            print(f"PASS: Component 1 -- horizontal separator found: {separator_desc} (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 -- no horizontal separator found after Subject line")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: The border/separator has a visible line style (0.3 points)
    # Must not be 'none' or 'nil' -- should be 'single', 'thick', 'double', etc.
    try:
        if separator_found:
            detected_style = _find_visible_border_style(doc, subject_idx)
            if detected_style:
                print(f"PASS: Component 2 -- border style is '{detected_style}' (visible) (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 -- no visible border style found")
        else:
            print(f"FAIL: Component 2 -- skipped (no separator found)")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Separator exists AND body text is preserved below it (0.2 points)
    # Compound check: separator must be present AND body text intact
    # This FAILS on initial (no separator) and PASSES on golden (separator + body)
    try:
        if separator_found:
            body_idx = _body_text_exists_after(doc, subject_idx, "I am writing to inform you")
            if body_idx is not None:
                print(f"PASS: Component 3 -- separator present AND body text preserved at P{body_idx} (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 -- separator found but body text missing")
        else:
            print(f"FAIL: Component 3 -- no separator found (prerequisite not met)")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
