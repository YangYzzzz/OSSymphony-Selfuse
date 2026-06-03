"""
Initial Setup: Memo document without horizontal separator
Task ID: writer_biz_007
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_007'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Memo Header Section ---
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run('MEMORANDUM')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Calibri'

    # Empty line after title
    doc.add_paragraph()

    # To line
    to_para = doc.add_paragraph()
    to_label = to_para.add_run('To:\t')
    to_label.bold = True
    to_label.font.size = Pt(11)
    to_value = to_para.add_run('All Department Managers')
    to_value.font.size = Pt(11)

    # From line
    from_para = doc.add_paragraph()
    from_label = from_para.add_run('From:\t')
    from_label.bold = True
    from_label.font.size = Pt(11)
    from_value = from_para.add_run('Patricia Nakamura, VP of Operations')
    from_value.font.size = Pt(11)

    # Date line
    date_para = doc.add_paragraph()
    date_label = date_para.add_run('Date:\t')
    date_label.bold = True
    date_label.font.size = Pt(11)
    date_value = date_para.add_run('March 28, 2026')
    date_value.font.size = Pt(11)

    # Subject line (NO bottom border - task is to add it)
    subject_para = doc.add_paragraph()
    subject_label = subject_para.add_run('Subject:\t')
    subject_label.bold = True
    subject_label.font.size = Pt(11)
    subject_value = subject_para.add_run('Q1 2026 Budget Review and Resource Allocation')
    subject_value.font.size = Pt(11)

    # --- Body (no separator - that's the task) ---
    doc.add_paragraph()

    body1 = doc.add_paragraph()
    body1_run = body1.add_run(
        'I am writing to inform you of the upcoming Q1 2026 budget review '
        'meeting scheduled for April 15, 2026. Each department manager is '
        'expected to prepare a comprehensive summary of their current spending '
        'versus allocated budget for the first quarter.'
    )
    body1_run.font.size = Pt(11)

    body2 = doc.add_paragraph()
    body2_run = body2.add_run(
        'As discussed in our last leadership meeting, the company has seen a '
        '12% increase in operational costs compared to the same period last year. '
        'Given this trend, we need to identify areas where we can optimize '
        'spending without compromising project deliverables or employee satisfaction.'
    )
    body2_run.font.size = Pt(11)

    body3 = doc.add_paragraph()
    body3_run = body3.add_run(
        'Please bring the following to the review meeting:'
    )
    body3_run.font.size = Pt(11)

    # Bullet items
    items = [
        'Detailed expense report for January through March 2026',
        'Year-over-year comparison with Q1 2025 figures',
        'Projected spending for Q2 2026 with justification',
        'Proposed cost-saving measures for the remainder of the fiscal year',
    ]
    for item in items:
        bullet = doc.add_paragraph(item, style='List Bullet')
        for run in bullet.runs:
            run.font.size = Pt(11)

    body4 = doc.add_paragraph()
    body4_run = body4.add_run(
        'If you have any questions or require additional data from the finance '
        'team, please reach out to David Romero at d.romero@westfield-corp.com '
        'before April 10, 2026.'
    )
    body4_run.font.size = Pt(11)

    closing = doc.add_paragraph()
    closing_run = closing.add_run('Thank you for your continued diligence and cooperation.')
    closing_run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
