"""
Initial Setup: Legal contract with 'Terms and Conditions' section containing 5 paragraphs
Task ID: osworld_writer_blank_line_insertion_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_blank_line_insertion_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("SERVICE AGREEMENT CONTRACT")
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph()

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run("Between Nexus Technologies LLC and Meridian Consulting Group")
    subtitle_run.font.size = Pt(12)
    subtitle_run.italic = True

    doc.add_paragraph()

    # --- Section 1: Preamble ---
    sec1_heading = doc.add_paragraph()
    sec1_run = sec1_heading.add_run("1. PREAMBLE")
    sec1_run.bold = True
    sec1_run.font.size = Pt(13)

    doc.add_paragraph(
        "This Service Agreement (\"Agreement\") is entered into as of January 15, 2025, "
        "by and between Nexus Technologies LLC, a limited liability company organized under the "
        "laws of the State of Delaware, with its principal place of business at 4200 Innovation "
        "Drive, Suite 800, Wilmington, DE 19801 (\"Service Provider\"), and Meridian Consulting "
        "Group, a corporation organized under the laws of the State of New York, with its "
        "principal place of business at 300 Park Avenue, 22nd Floor, New York, NY 10022 (\"Client\")."
    )

    doc.add_paragraph()

    # --- Section 2: Scope of Services ---
    sec2_heading = doc.add_paragraph()
    sec2_run = sec2_heading.add_run("2. SCOPE OF SERVICES")
    sec2_run.bold = True
    sec2_run.font.size = Pt(13)

    doc.add_paragraph(
        "The Service Provider agrees to deliver cloud infrastructure management, software "
        "development support, and cybersecurity consulting services as detailed in Exhibit A "
        "attached hereto and incorporated herein by reference. All deliverables shall meet the "
        "quality standards outlined in Exhibit B."
    )

    doc.add_paragraph(
        "Services shall commence on February 1, 2025, and shall continue for an initial term of "
        "twelve (12) months unless earlier terminated in accordance with Section 9 of this Agreement."
    )

    doc.add_paragraph()

    # --- Section 3: Payment Terms ---
    sec3_heading = doc.add_paragraph()
    sec3_run = sec3_heading.add_run("3. PAYMENT TERMS")
    sec3_run.bold = True
    sec3_run.font.size = Pt(13)

    doc.add_paragraph(
        "Client agrees to pay Service Provider a monthly retainer fee of $18,500 USD, due on the "
        "first business day of each calendar month. Late payments will accrue interest at a rate of "
        "1.5% per month on the outstanding balance."
    )

    doc.add_paragraph(
        "All invoices are payable within thirty (30) days of receipt. Service Provider reserves "
        "the right to suspend services for accounts overdue by more than forty-five (45) days."
    )

    doc.add_paragraph()

    # --- Section 4: Terms and Conditions ---
    sec4_heading = doc.add_paragraph()
    sec4_run = sec4_heading.add_run("4. TERMS AND CONDITIONS")
    sec4_run.bold = True
    sec4_run.font.size = Pt(13)

    # 5 paragraphs with single paragraph breaks only (no blank lines between them)
    doc.add_paragraph(
        "4.1 Confidentiality. Each party agrees to maintain in strict confidence all proprietary "
        "information, trade secrets, and business data disclosed by the other party in connection "
        "with this Agreement. Neither party shall disclose such confidential information to any "
        "third party without prior written consent, except as required by applicable law or court order. "
        "This obligation of confidentiality shall survive termination of this Agreement for a period "
        "of five (5) years."
    )

    doc.add_paragraph(
        "4.2 Intellectual Property. All work product, inventions, software code, and deliverables "
        "created by Service Provider solely for Client under this Agreement shall be considered "
        "work-for-hire and shall become the exclusive property of Client upon full payment of all "
        "fees due. Service Provider retains ownership of its pre-existing tools, methodologies, "
        "and general know-how."
    )

    doc.add_paragraph(
        "4.3 Limitation of Liability. In no event shall either party be liable to the other for any "
        "indirect, incidental, special, consequential, or punitive damages arising out of or related "
        "to this Agreement, even if such party has been advised of the possibility of such damages. "
        "The total cumulative liability of Service Provider under this Agreement shall not exceed the "
        "total fees paid by Client in the three (3) months preceding the event giving rise to the claim."
    )

    doc.add_paragraph(
        "4.4 Dispute Resolution. Any dispute, controversy, or claim arising out of or relating to "
        "this Agreement, or the breach, termination, or validity thereof, shall be finally resolved "
        "by binding arbitration in accordance with the Commercial Arbitration Rules of the American "
        "Arbitration Association. The arbitration shall be conducted in New York, New York, and the "
        "arbitrator's decision shall be final and binding upon both parties."
    )

    doc.add_paragraph(
        "4.5 Governing Law. This Agreement shall be governed by and construed in accordance with "
        "the laws of the State of New York, without giving effect to any choice-of-law or "
        "conflict-of-law rules. The parties consent to the exclusive jurisdiction of the courts "
        "located in New York County, New York, for any action arising out of or related to this "
        "Agreement that is not subject to arbitration."
    )

    doc.add_paragraph()

    # --- Section 5: Termination ---
    sec5_heading = doc.add_paragraph()
    sec5_run = sec5_heading.add_run("5. TERMINATION")
    sec5_run.bold = True
    sec5_run.font.size = Pt(13)

    doc.add_paragraph(
        "Either party may terminate this Agreement upon sixty (60) days' written notice to the other "
        "party. Service Provider may terminate immediately if Client fails to cure a material breach "
        "within fifteen (15) days of written notice thereof."
    )

    doc.add_paragraph()

    # --- Signature block ---
    sig_heading = doc.add_paragraph()
    sig_run = sig_heading.add_run("SIGNATURES")
    sig_run.bold = True
    sig_run.font.size = Pt(13)

    doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.")
    doc.add_paragraph()
    doc.add_paragraph("Nexus Technologies LLC                      Meridian Consulting Group")
    doc.add_paragraph()
    doc.add_paragraph("By: _______________________________         By: _______________________________")
    doc.add_paragraph("Name: Alexandra M. Foster                   Name: Robert J. Harrington")
    doc.add_paragraph("Title: Chief Executive Officer              Title: Managing Director")
    doc.add_paragraph("Date: ____________________________          Date: ____________________________")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
