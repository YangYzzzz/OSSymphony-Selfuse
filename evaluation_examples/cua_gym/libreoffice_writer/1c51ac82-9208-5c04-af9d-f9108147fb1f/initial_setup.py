"""
Initial Setup: Create API documentation with 4 parameter tables (default formatting)
Task ID: writer_tech_067
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_067'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_api_table(doc, headers, data):
    """Add a parameter table with default formatting (no custom styles)."""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    # Data rows
    for i, row_data in enumerate(data, 1):
        for j, val in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    return table


def create_initial():
    doc = Document()

    # Document title
    title = doc.add_heading('CloudVault API Reference', level=0)

    intro = doc.add_paragraph(
        'This document provides a comprehensive reference for the CloudVault REST API. '
        'Each endpoint is described with its parameters, expected types, and usage notes.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # ---- Endpoint 1: Authentication ----
    doc.add_heading('1. POST /api/v2/auth/login', level=1)
    desc1 = doc.add_paragraph(
        'Authenticates a user and returns an access token with configurable expiration. '
        'Supports multi-factor authentication when enabled on the account.'
    )
    desc1.paragraph_format.space_after = Pt(6)

    doc.add_paragraph('Request Parameters:', style='List Bullet')
    add_api_table(doc,
        ['Parameter', 'Type', 'Required', 'Description'],
        [
            ['username', 'string', 'Yes', 'The user\'s registered email address or username'],
            ['password', 'string', 'Yes', 'Account password (min 8 characters)'],
            ['mfa_code', 'string', 'No', 'Six-digit TOTP code if MFA is enabled'],
            ['remember_device', 'boolean', 'No', 'If true, extends token validity to 30 days'],
            ['client_id', 'string', 'Yes', 'OAuth2 application client identifier'],
        ]
    )
    doc.add_paragraph('')  # spacer

    # ---- Endpoint 2: File Upload ----
    doc.add_heading('2. POST /api/v2/files/upload', level=1)
    desc2 = doc.add_paragraph(
        'Uploads a file to the specified vault directory. Supports chunked uploads for '
        'files exceeding 100 MB. Returns the file metadata object on success.'
    )
    desc2.paragraph_format.space_after = Pt(6)

    doc.add_paragraph('Request Parameters:', style='List Bullet')
    add_api_table(doc,
        ['Parameter', 'Type', 'Required', 'Description'],
        [
            ['file', 'binary', 'Yes', 'The file content as multipart/form-data'],
            ['vault_id', 'string', 'Yes', 'Target vault UUID (36-char format)'],
            ['directory_path', 'string', 'No', 'Destination path within vault (default: root)'],
            ['overwrite', 'boolean', 'No', 'Replace existing file with same name if true'],
            ['encryption_key', 'string', 'No', 'Client-side AES-256 encryption key (base64)'],
            ['tags', 'array[string]', 'No', 'Metadata tags for search and categorization'],
        ]
    )
    doc.add_paragraph('')  # spacer

    # ---- Endpoint 3: Search ----
    doc.add_heading('3. GET /api/v2/search', level=1)
    desc3 = doc.add_paragraph(
        'Performs a full-text search across vault contents including file names, metadata, '
        'and OCR-extracted text from documents and images.'
    )
    desc3.paragraph_format.space_after = Pt(6)

    doc.add_paragraph('Query Parameters:', style='List Bullet')
    add_api_table(doc,
        ['Parameter', 'Type', 'Required', 'Description'],
        [
            ['query', 'string', 'Yes', 'Search query string (supports Lucene syntax)'],
            ['vault_ids', 'array[string]', 'No', 'Limit search to specific vault UUIDs'],
            ['file_types', 'array[string]', 'No', 'Filter by MIME type (e.g., application/pdf)'],
            ['date_from', 'ISO 8601', 'No', 'Include files modified after this timestamp'],
            ['date_to', 'ISO 8601', 'No', 'Include files modified before this timestamp'],
            ['page', 'integer', 'No', 'Page number for paginated results (default: 1)'],
            ['per_page', 'integer', 'No', 'Results per page, max 100 (default: 25)'],
        ]
    )
    doc.add_paragraph('')  # spacer

    # ---- Endpoint 4: Sharing ----
    doc.add_heading('4. POST /api/v2/share/create', level=1)
    desc4 = doc.add_paragraph(
        'Creates a shareable link for a file or directory. Supports expiration dates, '
        'password protection, and granular permission controls for recipients.'
    )
    desc4.paragraph_format.space_after = Pt(6)

    doc.add_paragraph('Request Parameters:', style='List Bullet')
    add_api_table(doc,
        ['Parameter', 'Type', 'Required', 'Description'],
        [
            ['resource_id', 'string', 'Yes', 'UUID of the file or directory to share'],
            ['resource_type', 'enum', 'Yes', 'Either "file" or "directory"'],
            ['permission', 'enum', 'Yes', 'Access level: "view", "download", or "edit"'],
            ['expires_at', 'ISO 8601', 'No', 'Link expiration timestamp (default: 7 days)'],
            ['password', 'string', 'No', 'Optional access password for the share link'],
            ['max_downloads', 'integer', 'No', 'Limit total downloads (0 = unlimited)'],
        ]
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
