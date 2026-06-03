"""
Initial Setup: Thesis document with chapter headings and empty header
Task ID: writer_acad_077
Domain: libreoffice_writer

Creates a multi-chapter thesis document with Heading 1 chapter numbering
configured. The header area exists but is empty — no Chapter field reference.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import lxml.etree as etree

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_077'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_chapter_numbering(doc):
    """
    Add outline numbering definition linked to Heading 1 style
    (equivalent to Tools > Chapter Numbering in LibreOffice).
    This creates an abstract numbering that associates level 0 with 'Heading 1'.
    """
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part.numbering_definitions._numbering

    # Create abstract numbering for chapter numbering
    abstract_num_xml = (
        '<w:abstractNum w:abstractNumId="10" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '  <w:multiLevelType w:val="multilevel"/>'
        '  <w:lvl w:ilvl="0">'
        '    <w:start w:val="1"/>'
        '    <w:numFmt w:val="decimal"/>'
        '    <w:pStyle w:val="Heading1"/>'
        '    <w:lvlText w:val="Chapter %1."/>'
        '    <w:lvlJc w:val="left"/>'
        '    <w:pPr><w:ind w:left="0" w:firstLine="0"/></w:pPr>'
        '  </w:lvl>'
        '</w:abstractNum>'
    )
    abstract_num = parse_xml(abstract_num_xml)
    numbering_elm.append(abstract_num)

    # Create concrete numbering referencing the abstract
    num_xml = (
        '<w:num w:numId="10" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '  <w:abstractNumId w:val="10"/>'
        '</w:num>'
    )
    num_elm = parse_xml(num_xml)
    numbering_elm.append(num_elm)

    return 10  # numId


def apply_heading1_numbering(para, num_id):
    """Apply numbering to a Heading 1 paragraph."""
    pPr = para._element.get_or_add_pPr()
    numPr = parse_xml(
        f'<w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'  <w:ilvl w:val="0"/>'
        f'  <w:numId w:val="{num_id}"/>'
        f'</w:numPr>'
    )
    pPr.append(numPr)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.25)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)

    # --- Configure empty header (exists but no content) ---
    header = section.header
    header.is_linked_to_previous = False
    # Clear any default paragraph text (keep the header area itself)
    for p in header.paragraphs:
        p.text = ""

    # --- Title Page ---
    title_para = doc.add_heading("Adaptive Machine Learning Approaches for\nClimate Prediction Models", level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run("by\nDr. Elena Vasquez-Moreno")
    run.font.size = Pt(14)

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = dept.add_run("Department of Computational Earth Sciences\nUniversity of Pacific Ridge\n\nA dissertation submitted in partial fulfillment\nof the requirements for the degree of\nDoctor of Philosophy\n\nMarch 2026")
    run.font.size = Pt(12)

    doc.add_page_break()

    # --- Add chapter numbering ---
    num_id = add_chapter_numbering(doc)

    # === Chapter 1: Introduction ===
    ch1 = doc.add_heading("Introduction", level=1)
    apply_heading1_numbering(ch1, num_id)

    doc.add_paragraph(
        "The increasing frequency of extreme weather events has underscored the urgent need "
        "for more accurate and timely climate prediction models. Traditional numerical weather "
        "prediction (NWP) systems, while grounded in well-established physical equations, often "
        "struggle to capture the nonlinear dynamics inherent in complex climate systems. This "
        "limitation has spurred a growing interest in machine learning (ML) techniques as either "
        "complementary or alternative approaches to conventional modeling frameworks."
    )
    doc.add_paragraph(
        "Over the past decade, deep learning architectures — particularly convolutional neural "
        "networks (CNNs) and recurrent neural networks (RNNs) — have demonstrated remarkable "
        "success in pattern recognition tasks across various scientific domains. In atmospheric "
        "science, these models have been applied to precipitation nowcasting, tropical cyclone "
        "intensity estimation, and sub-seasonal to seasonal (S2S) forecasting with promising "
        "results. However, significant challenges remain regarding model interpretability, "
        "physical consistency, and generalization to unseen climate regimes."
    )
    doc.add_paragraph(
        "This dissertation addresses these challenges by proposing an adaptive hybrid framework "
        "that integrates physics-informed neural networks (PINNs) with ensemble Kalman filter "
        "techniques. The central hypothesis posits that embedding physical conservation laws "
        "directly into the learning architecture yields predictions that are both more accurate "
        "and more physically plausible than purely data-driven approaches."
    )

    doc.add_page_break()

    # === Chapter 2: Literature Review ===
    ch2 = doc.add_heading("Literature Review", level=1)
    apply_heading1_numbering(ch2, num_id)

    doc.add_heading("Traditional Climate Modeling", level=2)
    doc.add_paragraph(
        "General Circulation Models (GCMs) have served as the backbone of climate science "
        "since the pioneering work of Manabe and Wetherald (1967). These models discretize "
        "the Earth's atmosphere into three-dimensional grid cells and solve the primitive "
        "equations governing atmospheric motion, thermodynamics, and radiative transfer. "
        "Despite continuous improvements in spatial resolution and parameterization schemes, "
        "GCMs remain computationally expensive and are subject to systematic biases arising "
        "from unresolved sub-grid-scale processes."
    )

    doc.add_heading("Machine Learning in Geosciences", level=2)
    doc.add_paragraph(
        "The application of ML to geoscientific problems dates back to the early 1990s when "
        "simple feedforward networks were used for empirical downscaling of GCM outputs. "
        "Rasp et al. (2018) demonstrated that deep learning could serve as a surrogate for "
        "computationally expensive convective parameterization schemes, achieving comparable "
        "accuracy at a fraction of the computational cost. More recently, Weyn et al. (2020) "
        "showed that graph neural networks trained on ERA5 reanalysis data could produce "
        "skillful medium-range weather forecasts competitive with operational NWP models."
    )

    doc.add_heading("Physics-Informed Neural Networks", level=2)
    doc.add_paragraph(
        "Raissi et al. (2019) introduced PINNs as a framework for embedding differential "
        "equation constraints directly into the loss function of neural networks. This "
        "approach has since been extended to a wide range of fluid dynamics problems, "
        "including Navier-Stokes equations, heat transfer, and turbulence modeling. In the "
        "context of climate science, Beucler et al. (2021) demonstrated that enforcing "
        "energy conservation constraints during training improved the physical plausibility "
        "of neural network-based climate emulators."
    )

    doc.add_page_break()

    # === Chapter 3: Methodology ===
    ch3 = doc.add_heading("Methodology", level=1)
    apply_heading1_numbering(ch3, num_id)

    doc.add_heading("Data Sources and Preprocessing", level=2)
    doc.add_paragraph(
        "This study utilizes ERA5 reanalysis data from the European Centre for Medium-Range "
        "Weather Forecasts (ECMWF) covering the period 1979-2024. The dataset includes "
        "hourly atmospheric variables at 37 pressure levels with a horizontal resolution of "
        "0.25 degrees. Key variables include temperature, specific humidity, geopotential "
        "height, and wind components (u, v). Additional sea surface temperature data from "
        "NOAA's OISST v2.1 product are incorporated to capture ocean-atmosphere coupling."
    )

    doc.add_heading("Hybrid Architecture Design", level=2)
    doc.add_paragraph(
        "The proposed Adaptive Physics-Informed Climate Network (APICN) consists of three "
        "main components: (1) a spatiotemporal encoder based on 3D convolutional layers that "
        "extracts multi-scale atmospheric features, (2) a physics constraint module that "
        "enforces conservation of mass, energy, and momentum through Lagrangian penalties, "
        "and (3) an ensemble prediction head that generates probabilistic forecasts via "
        "variational dropout. The architecture is trained end-to-end using a composite loss "
        "function that balances predictive accuracy against physical constraint violations."
    )

    doc.add_heading("Ensemble Kalman Filter Integration", level=2)
    doc.add_paragraph(
        "To address the challenge of sequential data assimilation, we integrate an Ensemble "
        "Kalman Filter (EnKF) into the prediction pipeline. At each forecast step, the EnKF "
        "updates the neural network's latent state representation using incoming observational "
        "data, effectively combining the learned climate dynamics with real-time measurements. "
        "This hybrid approach enables the model to adapt to evolving atmospheric conditions "
        "and correct for systematic biases that may accumulate during extended forecast horizons."
    )

    doc.add_page_break()

    # === Chapter 4: Results and Discussion ===
    ch4 = doc.add_heading("Results and Discussion", level=1)
    apply_heading1_numbering(ch4, num_id)

    doc.add_heading("Forecast Skill Evaluation", level=2)
    doc.add_paragraph(
        "The APICN model was evaluated against three benchmark systems: the operational "
        "ECMWF IFS model, a standard ResNet-based weather prediction model, and a purely "
        "data-driven transformer architecture (FourCastNet). Evaluation metrics include the "
        "anomaly correlation coefficient (ACC), root mean square error (RMSE), and the "
        "continuous ranked probability score (CRPS) for probabilistic assessments."
    )
    doc.add_paragraph(
        "For 500 hPa geopotential height predictions, APICN achieved an ACC of 0.89 at "
        "day 7, compared to 0.91 for ECMWF IFS, 0.84 for ResNet, and 0.86 for FourCastNet. "
        "At extended ranges (day 10-14), the physics-informed approach showed a notable "
        "advantage, maintaining an ACC above 0.72 while the purely data-driven models "
        "deteriorated to below 0.65. This suggests that the embedded physical constraints "
        "help prevent the accumulation of non-physical artifacts during long-range forecasts."
    )

    doc.add_heading("Physical Consistency Analysis", level=2)
    doc.add_paragraph(
        "A critical advantage of the APICN framework is the improved physical consistency "
        "of its predictions. Energy budget analysis revealed that APICN predictions conserve "
        "total atmospheric energy to within 0.3 W/m^2, compared to 2.1 W/m^2 for the "
        "unconstrained ResNet and 1.4 W/m^2 for FourCastNet. Similarly, moisture conservation "
        "errors were reduced by approximately 60% relative to purely data-driven baselines."
    )

    doc.add_page_break()

    # === Chapter 5: Conclusion ===
    ch5 = doc.add_heading("Conclusion", level=1)
    apply_heading1_numbering(ch5, num_id)

    doc.add_paragraph(
        "This dissertation has presented the Adaptive Physics-Informed Climate Network, a "
        "novel hybrid framework that integrates physics-based constraints with deep learning "
        "for improved climate prediction. The key findings demonstrate that embedding physical "
        "conservation laws into the neural network architecture yields predictions that are "
        "both competitive with state-of-the-art operational models and significantly more "
        "physically consistent than purely data-driven approaches."
    )
    doc.add_paragraph(
        "The integration of the Ensemble Kalman Filter for sequential data assimilation "
        "represents a particularly promising avenue for operational deployment, as it enables "
        "the model to continuously adapt to incoming observations while maintaining physical "
        "plausibility. Future work should explore the extension of this framework to "
        "higher-resolution regional modeling and the incorporation of additional physical "
        "constraints related to cloud microphysics and aerosol-radiation interactions."
    )

    # --- Save ---
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # --- GUI Launch ---
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
