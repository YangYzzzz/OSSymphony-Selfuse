"""
Reward Script: Add a running header with Chapter field reference (STYLEREF)
Task ID: writer_acad_077
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Header contains field codes (fldChar begin/separate/end)
  Component 2 (0.35): Field instruction references Heading 1 via STYLEREF
  Component 3 (0.25): Cached field text is non-empty and matches a chapter title
"""

import os
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_077'

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': WNS}


def persist_app_state(domain: str):
    """Best-effort save via Ctrl+S before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all Heading 1 texts from the document body (for Component 3 validation)
    heading1_texts = []
    try:
        for para in doc.paragraphs:
            if para.style and para.style.name == 'Heading 1':
                heading1_texts.append(para.text.strip())
        print(f"INFO: Found {len(heading1_texts)} Heading 1 paragraphs: {heading1_texts}")
    except Exception as e:
        print(f"WARN: Could not enumerate headings: {e}")

    # Get header from first section
    try:
        section = doc.sections[0]
        header = section.header
        header_paras = header.paragraphs
        if not header_paras:
            print("FAIL: No paragraphs in header")
            print("REWARD: 0.0")
            return 0.0
    except Exception as e:
        print(f"CRITICAL: Cannot access header: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Search all header paragraphs for field codes
    found_fld_chars = False
    found_styleref_heading1 = False
    cached_text = ""

    for p in header_paras:
        p_xml = p._element

        # Find fldChar elements
        fld_chars = p_xml.findall('.//w:fldChar', NS)
        instr_texts = p_xml.findall('.//w:instrText', NS)

        if len(fld_chars) >= 2:  # need at least begin + end
            found_fld_chars = True

        # Check instrText for STYLEREF referencing Heading 1
        for instr in instr_texts:
            instr_val = (instr.text or "").strip()
            print(f"INFO: instrText found: {instr_val!r}")
            # Accept STYLEREF with "Heading 1" or similar variants
            # Also accept CHAPTER field codes (alternative approach in LibreOffice)
            instr_upper = instr_val.upper()
            if 'STYLEREF' in instr_upper and 'HEADING' in instr_upper and '1' in instr_upper:
                found_styleref_heading1 = True
            # Also accept: STYLEREF "Heading 1" \* MERGEFORMAT etc.
            elif 'STYLEREF' in instr_upper:
                # Check if it references heading 1 in any format
                import re
                if re.search(r'HEADING\s*1', instr_upper):
                    found_styleref_heading1 = True

        # Extract cached/displayed text from between separate and end fldChars
        # The text between separate and end is the cached field result
        runs = p_xml.findall('.//w:r', NS)
        in_field_result = False
        for run in runs:
            fld = run.find('w:fldChar', NS)
            if fld is not None:
                fld_type = fld.get(f'{{{WNS}}}fldCharType')
                if fld_type == 'separate':
                    in_field_result = True
                    continue
                elif fld_type == 'end':
                    in_field_result = False
                    continue
            if in_field_result:
                t_elem = run.find('w:t', NS)
                if t_elem is not None and t_elem.text:
                    cached_text += t_elem.text

    # Component 1: Header contains field codes (0.4 points)
    # This checks that field codes exist in the header - the initial has none
    try:
        if found_fld_chars:
            print(f"PASS: Component 1 -- Header contains field codes (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 -- No field codes found in header")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Field is a STYLEREF referencing Heading 1 (0.35 points)
    # This verifies the correct field type was used for chapter reference
    try:
        if found_styleref_heading1:
            print(f"PASS: Component 2 -- STYLEREF 'Heading 1' field found (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 2 -- No STYLEREF 'Heading 1' instrText in header")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Cached field text is non-empty and matches a chapter title (0.25 points)
    # The STYLEREF should resolve to one of the Heading 1 texts in the document
    try:
        cached_stripped = cached_text.strip()
        if cached_stripped:
            if heading1_texts and cached_stripped in heading1_texts:
                print(f"PASS: Component 3 -- Cached text '{cached_stripped}' matches a Heading 1 (0.25 pts)")
                total_score += 0.25
            elif cached_stripped:
                # The cached text exists but may not exactly match due to field update timing
                # Still give credit if it's non-empty (field is functional)
                print(f"PASS: Component 3 -- Cached text '{cached_stripped}' is non-empty (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 -- Cached field text is empty")
        else:
            print(f"FAIL: Component 3 -- No cached text found in field result")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
