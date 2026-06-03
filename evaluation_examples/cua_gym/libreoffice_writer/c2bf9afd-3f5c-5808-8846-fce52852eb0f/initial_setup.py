"""
Initial Setup: Q4 2025 Performance Report - Initial State
Task ID: writer_biz_018
Domain: libreoffice_writer

Creates report_v3_final_FINAL.docx on ~/Desktop/ with 6 pages of quarterly
performance data but NO title page. The agent must rename the file and add
a title page.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_018'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/report_v3_final_FINAL.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set default styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # -------------------------------------------------------
    # Page 1 — Introduction
    # -------------------------------------------------------
    h = doc.add_heading('Introduction', level=1)
    h.paragraph_format.space_before = Pt(0)

    intro_para = doc.add_paragraph(
        'This report presents the quarterly performance results for Q4 2025 '
        'for Crestview Holdings and its subsidiaries. The data herein covers '
        'financial metrics, operational KPIs, and strategic highlights for '
        'the period ending December 31, 2025. All figures are preliminary '
        'and subject to audit adjustment.'
    )
    intro_para.paragraph_format.space_after = Pt(8)

    doc.add_paragraph(
        'The analysis encompasses four business units: Corporate Banking, '
        'Retail Financial Services, Asset Management, and Insurance. Each '
        'unit\'s performance is evaluated against Q3 2025 actuals and '
        'FY 2025 targets established at the beginning of the fiscal year.'
    )

    doc.add_paragraph(
        'Key macro-economic factors influencing Q4 2025 results include '
        'rising interest rates, continued inflationary pressure in operational '
        'costs, and a modest recovery in equity markets during November and '
        'December. Management commentary accompanying each section provides '
        'context for variance against forecast.'
    )

    doc.add_page_break()

    # -------------------------------------------------------
    # Page 2 — Financial Summary
    # -------------------------------------------------------
    doc.add_heading('Financial Summary', level=1)

    doc.add_paragraph(
        'Consolidated revenue for Q4 2025 reached $312.4 million, representing '
        'a 7.2% increase year-over-year and a 3.1% sequential increase from Q3 2025. '
        'Operating income was $78.9 million, reflecting an operating margin of 25.3%. '
        'Net income attributable to shareholders totalled $54.2 million.'
    )

    # Financial table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ['Metric', 'Q4 2025', 'Q3 2025', 'Q4 2024']
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True

    data_rows = [
        ['Total Revenue ($M)', '312.4', '303.0', '291.5'],
        ['Operating Income ($M)', '78.9', '74.1', '69.3'],
        ['Net Income ($M)', '54.2', '49.8', '45.7'],
        ['EPS (diluted)', '$2.31', '$2.12', '$1.95'],
        ['EBITDA ($M)', '94.7', '89.2', '83.4'],
    ]
    for row_idx, row_data in enumerate(data_rows, 1):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx, col_idx).text = value

    doc.add_page_break()

    # -------------------------------------------------------
    # Page 3 — Corporate Banking Performance
    # -------------------------------------------------------
    doc.add_heading('Corporate Banking Performance', level=1)

    doc.add_paragraph(
        'The Corporate Banking division delivered $128.6 million in revenue during '
        'Q4 2025, driven by strong loan origination volumes and improved net interest '
        'margins. Total loan portfolio grew to $2.1 billion, a 4.8% increase from '
        'Q3 2025. Non-performing loans as a percentage of total loans improved to '
        '1.2%, down from 1.5% in the prior quarter.'
    )

    doc.add_heading('Key Highlights', level=2)
    bullets = [
        'New corporate mandates: 14 deals closed totaling $340M in commitments',
        'Trade finance volumes up 18% year-over-year to $156M',
        'Average deposit balance grew 6.3% to $1.4 billion',
        'Cost-to-income ratio improved to 42.1% from 44.8% in Q3 2025',
        'Treasury management revenue: $18.3M (up $2.1M from prior quarter)',
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    doc.add_paragraph(
        'Credit quality remained resilient despite tighter lending conditions. '
        'The division maintained a disciplined underwriting approach with '
        'weighted average loan-to-value ratios at 58% across the commercial '
        'real estate portfolio. Provisions for credit losses were $4.2 million, '
        'consistent with Q3 2025 levels.'
    )

    doc.add_page_break()

    # -------------------------------------------------------
    # Page 4 — Retail Financial Services
    # -------------------------------------------------------
    doc.add_heading('Retail Financial Services', level=1)

    doc.add_paragraph(
        'Retail Financial Services generated $89.4 million in Q4 2025, reflecting '
        'growth in mortgage originations and improved cross-sell ratios. Net new '
        'accounts opened totalled 12,847 during the quarter, bringing total active '
        'retail customers to 487,300.'
    )

    doc.add_heading('Mortgage Portfolio', level=2)
    doc.add_paragraph(
        'Residential mortgage originations reached $654 million in Q4 2025, up 11% '
        'from Q3 2025. The proportion of fixed-rate mortgages increased to 67% of '
        'new originations as customers sought rate certainty. Refinancing activity '
        'accounted for 28% of total originations. Average loan size was $427,000.'
    )

    doc.add_heading('Digital Banking', level=2)
    doc.add_paragraph(
        'Digital channel adoption continued to accelerate with mobile banking active '
        'users reaching 234,500, a 15% increase year-over-year. Digital transaction '
        'volume represented 74% of all retail transactions. The new mobile app '
        'launched in October 2025 received a 4.6/5.0 customer satisfaction rating '
        'in post-launch surveys.'
    )

    doc.add_page_break()

    # -------------------------------------------------------
    # Page 5 — Asset Management & Insurance
    # -------------------------------------------------------
    doc.add_heading('Asset Management & Insurance', level=1)

    doc.add_heading('Asset Management', level=2)
    doc.add_paragraph(
        'Assets under management reached $6.8 billion as of December 31, 2025, '
        'up 8.4% from September 30, 2025. Net new inflows of $340 million were '
        'recorded in Q4 2025, with equity funds accounting for $210 million of '
        'the total. Fee revenue from asset management activities was $34.2 million.'
    )

    am_table = doc.add_table(rows=5, cols=3)
    am_table.style = 'Table Grid'
    am_headers = ['Fund Category', 'AUM ($B)', 'Net Flow ($M)']
    for col_idx, h_text in enumerate(am_headers):
        cell = am_table.cell(0, col_idx)
        run = cell.paragraphs[0].add_run(h_text)
        run.bold = True

    am_data = [
        ['Equity Funds', '2.84', '+210'],
        ['Fixed Income', '2.12', '+85'],
        ['Balanced Funds', '1.24', '+31'],
        ['Money Market', '0.60', '+14'],
    ]
    for row_idx, row_data in enumerate(am_data, 1):
        for col_idx, value in enumerate(row_data):
            am_table.cell(row_idx, col_idx).text = value

    doc.add_paragraph('')

    doc.add_heading('Insurance Division', level=2)
    doc.add_paragraph(
        'Insurance premiums written in Q4 2025 totalled $60.4 million, up 5.1% '
        'year-over-year. Combined ratio improved to 93.8% from 95.2% in Q4 2024, '
        'reflecting better underwriting discipline and lower claims frequency in '
        'the property segment. Life insurance in-force grew 3.2% to $4.9 billion.'
    )

    doc.add_page_break()

    # -------------------------------------------------------
    # Page 6 — Outlook & Closing Remarks
    # -------------------------------------------------------
    doc.add_heading('Outlook & Closing Remarks', level=1)

    doc.add_paragraph(
        'Looking ahead to Q1 2026, management remains cautiously optimistic. '
        'Revenue guidance for Q1 2026 is set in the range of $305–$320 million, '
        'reflecting seasonal patterns in retail banking activity and anticipated '
        'headwinds from continued rate normalization in fixed income markets.'
    )

    doc.add_heading('Strategic Priorities for 2026', level=2)
    priorities = [
        'Expand corporate banking market share in the mid-market segment',
        'Launch enhanced digital wealth management platform by Q2 2026',
        'Achieve cost efficiency ratio below 41% across all business units',
        'Grow insurance premium volumes by 8% through targeted product launches',
        'Strengthen ESG reporting framework to align with IFRS S1 and S2 standards',
    ]
    for priority in priorities:
        doc.add_paragraph(priority, style='List Bullet')

    doc.add_paragraph(
        'The Finance Department extends its appreciation to all business unit heads '
        'and their teams for their commitment to delivering accurate and timely '
        'performance data for this report. Any queries regarding the contents of '
        'this report should be directed to the Group Financial Controller.'
    )

    closing_para = doc.add_paragraph(
        'Confidential — For Internal Use Only. © 2025 Crestview Holdings. '
        'All rights reserved.'
    )
    closing_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in closing_para.runs:
        run.font.size = Pt(9)
        run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the initial file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
