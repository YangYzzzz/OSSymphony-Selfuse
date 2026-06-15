"""
Reward Script: Rename file and add title page to Writer document
Task ID: writer_biz_018
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Old file removed and new filename exists on Desktop
  Component 2 (0.25): Title paragraph 'Q4 2025 Performance Report' is present, centered, and bold
  Component 3 (0.20): 'Crestview Holdings' company name paragraph is present and centered
  Component 4 (0.15): 'Prepared by: Finance Department' paragraph is present
  Component 5 (0.15): Page break separates title page from original content (Introduction heading follows)
Total: 1.0
"""

import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

DESKTOP = '/home/user/Desktop'
OLD_FILENAME = 'report_v3_final_FINAL.docx'
NEW_FILENAME = 'Q4_2025_Performance_Report.docx'
OLD_FILE_PATH = os.path.join(DESKTOP, OLD_FILENAME)
NEW_FILE_PATH = os.path.join(DESKTOP, NEW_FILENAME)

TASK_ID = 'writer_biz_018'


def count_page_breaks(doc):
    """Count page break elements in the document."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    count = 0
    break_paras = []
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            for br in run.element.findall('.//w:br', ns):
                if br.attrib.get(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type'
                ) == 'page':
                    count += 1
                    break_paras.append(i)
    return count, break_paras


def verify_task():
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # -------------------------------------------------------------------------
    # Component 1: File rename (0.25 points)
    # Old filename must NOT exist; new filename MUST exist.
    # This is a task-introduced change (initial has old name, golden has new name).
    # -------------------------------------------------------------------------
    try:
        old_exists = os.path.exists(OLD_FILE_PATH)
        new_exists = os.path.exists(NEW_FILE_PATH)

        if not old_exists and new_exists:
            print(f"PASS: Component 1 — File renamed correctly. '{OLD_FILENAME}' absent, '{NEW_FILENAME}' present (0.25 pts)")
            total_score += 0.25
        elif not new_exists:
            print(f"FAIL: Component 1 — New file '{NEW_FILENAME}' not found on Desktop")
        elif old_exists:
            print(f"FAIL: Component 1 — Old file '{OLD_FILENAME}' still present on Desktop")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Load the document — precondition gate
    # If the file cannot be loaded, skip all remaining checks.
    # -------------------------------------------------------------------------
    if not os.path.exists(NEW_FILE_PATH):
        print(f"CRITICAL: Cannot open document — '{NEW_FILE_PATH}' not found")
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    try:
        doc = Document(NEW_FILE_PATH)
    except Exception as e:
        print(f"CRITICAL: Cannot load document '{NEW_FILE_PATH}': {e}")
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # -------------------------------------------------------------------------
    # Component 2: Title paragraph present, centered, and bold (0.25 points)
    # "Q4 2025 Performance Report" must appear early in the document,
    # be centered, and have at least one bold run.
    # -------------------------------------------------------------------------
    try:
        title_text = "Q4 2025 Performance Report"
        title_found = False
        title_para_idx = -1

        for i, para in enumerate(doc.paragraphs[:30]):
            if title_text.lower() in para.text.lower():
                title_found = True
                title_para_idx = i
                is_centered = (para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
                has_bold_run = any(
                    run.bold is True or run.font.bold is True
                    for run in para.runs
                    if run.text.strip()
                )
                if is_centered and has_bold_run:
                    print(f"PASS: Component 2 — Title '{title_text}' found at para {i}, centered and bold (0.25 pts)")
                    total_score += 0.25
                elif not is_centered:
                    print(f"FAIL: Component 2 — Title found at para {i} but NOT centered (alignment={para.paragraph_format.alignment})")
                elif not has_bold_run:
                    print(f"FAIL: Component 2 — Title found at para {i} but NOT bold")
                break

        if not title_found:
            print(f"FAIL: Component 2 — Title text '{title_text}' not found in first 30 paragraphs")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Company name "Crestview Holdings" present and centered (0.20 points)
    # Must appear on the title page, early in the document, centered.
    # -------------------------------------------------------------------------
    try:
        company_text = "Crestview Holdings"
        company_found = False
        # Search within the title page area (before the page break / before Introduction)
        for i, para in enumerate(doc.paragraphs[:30]):
            if company_text.lower() in para.text.lower():
                company_found = True
                is_centered = (para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
                if is_centered:
                    print(f"PASS: Component 3 — '{company_text}' found at para {i}, centered (0.20 pts)")
                    total_score += 0.20
                else:
                    print(f"FAIL: Component 3 — '{company_text}' found at para {i} but NOT centered (alignment={para.paragraph_format.alignment})")
                break

        if not company_found:
            print(f"FAIL: Component 3 — '{company_text}' not found in first 30 paragraphs")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: "Prepared by: Finance Department" present (0.15 points)
    # Must appear somewhere on the title page.
    # -------------------------------------------------------------------------
    try:
        prepared_text = "Prepared by: Finance Department"
        prepared_found = False

        for i, para in enumerate(doc.paragraphs[:35]):
            if prepared_text.lower() in para.text.lower():
                prepared_found = True
                print(f"PASS: Component 4 — '{prepared_text}' found at para {i} (0.15 pts)")
                total_score += 0.15
                break

        if not prepared_found:
            print(f"FAIL: Component 4 — '{prepared_text}' not found in first 35 paragraphs")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -------------------------------------------------------------------------
    # Component 5: Page break separates title page from original content (0.15 points)
    # There must be at least one page break before the "Introduction" heading.
    # The original document's first heading ("Introduction") must still be present.
    # -------------------------------------------------------------------------
    try:
        pb_count, pb_paras = count_page_breaks(doc)

        # Find the first occurrence of "Introduction" heading
        intro_para_idx = -1
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == 'Introduction' and 'Heading' in para.style.name:
                intro_para_idx = i
                break

        if intro_para_idx == -1:
            # Fallback: find "Introduction" anywhere
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip() == 'Introduction':
                    intro_para_idx = i
                    break

        if intro_para_idx > 0 and pb_count > 0:
            # Check if any page break occurs before the Introduction paragraph
            breaks_before_intro = [pb for pb in pb_paras if pb < intro_para_idx]
            if breaks_before_intro:
                print(f"PASS: Component 5 — Page break(s) found before 'Introduction' (para {intro_para_idx}); breaks at paras {breaks_before_intro} (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 5 — Page break exists but not before 'Introduction' (para {intro_para_idx}); break paras={pb_paras}")
        elif intro_para_idx == -1:
            print(f"FAIL: Component 5 — 'Introduction' paragraph not found in document (original content may be missing)")
        else:
            print(f"FAIL: Component 5 — No page break found in document (pb_count={pb_count})")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -------------------------------------------------------------------------
    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


verify_task()
