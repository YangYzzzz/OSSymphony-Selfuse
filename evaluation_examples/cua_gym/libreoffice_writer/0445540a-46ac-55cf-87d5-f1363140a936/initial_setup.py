"""
Initial Setup: Leave policy document with 6 section headings at 12pt regular weight
Task ID: writer_hr_013
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_013'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

HEADINGS = [
    "Purpose and Scope",
    "Types of Leave",
    "Leave Accrual and Carryover",
    "Request and Approval Process",
    "Leave During Probation Period",
    "Return to Work Policy",
]

BODY_TEXTS = {
    "Purpose and Scope": (
        "This policy outlines the leave entitlements and procedures for all full-time and part-time "
        "employees of Meridian Technologies Inc. It applies to staff across all departments and office "
        "locations, effective from January 1, 2025. The purpose is to ensure consistent and fair "
        "management of employee absences while maintaining operational efficiency. Temporary contractors "
        "and consultants should refer to their individual agreements for leave provisions."
    ),
    "Types of Leave": (
        "Employees are entitled to the following categories of leave:\n"
        "Annual Leave: Full-time employees accrue 20 working days per calendar year. Part-time employees "
        "receive a pro-rata allocation based on their contracted hours.\n"
        "Sick Leave: Up to 10 paid sick days per year. Medical certification is required for absences "
        "exceeding three consecutive days.\n"
        "Personal Leave: 3 days per year for personal matters including family emergencies and "
        "bereavement.\n"
        "Parental Leave: 16 weeks of paid maternity leave and 4 weeks of paid paternity leave, in "
        "accordance with local labor regulations.\n"
        "Unpaid Leave: May be granted at management discretion for periods up to 30 calendar days."
    ),
    "Leave Accrual and Carryover": (
        "Annual leave accrues on a monthly basis at the rate of 1.67 days per month for full-time "
        "employees. Unused annual leave may be carried over to the following year, up to a maximum of "
        "5 days. Any carried-over leave must be used by March 31 of the new year. Sick leave does not "
        "accrue and resets at the start of each calendar year. Employees who join or leave mid-year "
        "will have their entitlements calculated on a pro-rata basis from their start or end date."
    ),
    "Request and Approval Process": (
        "All leave requests must be submitted through the HR Portal at least 5 business days in advance "
        "for planned absences. Emergency or sick leave should be reported to the direct supervisor by "
        "9:00 AM on the first day of absence, followed by a formal request in the system within 48 hours. "
        "Approval authority rests with the employee's direct manager for requests of up to 5 consecutive "
        "days. Requests exceeding 5 days require additional approval from the Department Head. The HR "
        "department reserves the right to request supporting documentation for any leave type."
    ),
    "Leave During Probation Period": (
        "Employees in their probationary period (first 90 days) are entitled to sick leave only, at "
        "a reduced rate of 5 days. Annual leave begins accruing from day one but cannot be taken until "
        "the probation period is successfully completed. In exceptional circumstances, the HR Director "
        "may authorize early use of annual leave on a case-by-case basis. Personal leave is not available "
        "during probation."
    ),
    "Return to Work Policy": (
        "Employees returning from leave of 5 or more consecutive working days must schedule a return-to-work "
        "meeting with their manager on the first day back. For medical leave exceeding 10 days, a fitness-for-duty "
        "certificate from a licensed physician is required before resuming duties. Phased return arrangements "
        "may be made in coordination with HR and Occupational Health for employees recovering from serious "
        "illness or injury. All return-to-work documentation must be filed with the HR department within "
        "3 business days of the employee's return."
    ),
}


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Configure Heading 1 style: 12pt, NOT bold (regular weight)
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(12)
    h1_style.font.bold = False
    h1_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Title
    title = doc.add_heading("Meridian Technologies Inc.", level=0)
    subtitle = doc.add_paragraph("Employee Leave Policy 2025")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.italic = True

    doc.add_paragraph("")  # spacer

    # Add each section heading + body
    for heading_text in HEADINGS:
        h = doc.add_heading(heading_text, level=1)
        # Ensure the heading run is explicitly 12pt and not bold
        for run in h.runs:
            run.font.size = Pt(12)
            run.bold = False

        body = BODY_TEXTS[heading_text]
        doc.add_paragraph(body)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
