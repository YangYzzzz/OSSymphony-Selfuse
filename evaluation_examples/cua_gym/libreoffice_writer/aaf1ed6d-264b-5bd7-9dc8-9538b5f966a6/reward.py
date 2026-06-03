"""
Reward Script: Apply Heading 1 style to 'MOTION TO COMPEL DISCOVERY' title
Task ID: writer_legal_007
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): First paragraph has 'Heading 1' style applied
  Component 2 (0.4): First paragraph has 'Heading 1' style AND text preserved AND
                      other paragraphs not inadvertently changed to Heading 1
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_007'


def persist_app_state():
    """Save any unsaved LibreOffice edits before verification."""
    try:
        os.environ["DISPLAY"] = ":0"
        import pyautogui
        import time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 1 paragraph
    if len(doc.paragraphs) == 0:
        print("FAIL: Document has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    first_para = doc.paragraphs[0]
    first_style = first_para.style.name if first_para.style else None

    # Component 1: First paragraph has 'Heading 1' style (0.6 points)
    # This is the core task requirement. Initial has 'Normal', golden has 'Heading 1'.
    try:
        if first_style == 'Heading 1':
            print(f"PASS: Component 1 -- First paragraph style is 'Heading 1' (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 1 -- Expected style 'Heading 1', found '{first_style}'")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: First paragraph has 'Heading 1' AND text preserved AND
    #              other paragraphs not inadvertently changed to Heading 1 (0.4 points)
    # This compound check ensures the change was precise and correct.
    try:
        if first_style == 'Heading 1':
            # Sub-check A: text content preserved
            first_text = first_para.text.strip()
            text_ok = (first_text == 'MOTION TO COMPEL DISCOVERY')

            # Sub-check B: other paragraphs not changed to Heading 1
            other_heading1_count = 0
            for para in doc.paragraphs[1:]:
                if para.style and para.style.name == 'Heading 1':
                    other_heading1_count += 1
            no_extra_heading1 = (other_heading1_count == 0)

            if text_ok and no_extra_heading1:
                print(f"PASS: Component 2 -- Text preserved and no extra Heading 1 styles (0.4 pts)")
                total_score += 0.4
            else:
                if not text_ok:
                    print(f"FAIL: Component 2 -- Title text changed to '{first_text}'")
                if not no_extra_heading1:
                    print(f"FAIL: Component 2 -- {other_heading1_count} other paragraphs changed to Heading 1")
        else:
            print(f"FAIL: Component 2 -- First paragraph is not Heading 1, skipping compound check")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state()
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
