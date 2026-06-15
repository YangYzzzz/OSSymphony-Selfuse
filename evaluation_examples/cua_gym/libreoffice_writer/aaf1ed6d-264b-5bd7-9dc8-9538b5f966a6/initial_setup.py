"""
Initial Setup: Apply Heading 1 style to court filing title
Task ID: writer_legal_007
Domain: libreoffice_writer

Creates a Writer document with 'MOTION TO COMPEL DISCOVERY' as the first line
in Default Paragraph Style with manual bold and centering. Body text follows.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_007'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default page margins for a legal document
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Title line: Default Paragraph Style with manual bold + center ---
    title_para = doc.add_paragraph()
    title_para.style = doc.styles['Normal']  # Default Paragraph Style
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(12)
    title_run = title_para.add_run('MOTION TO COMPEL DISCOVERY')
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Times New Roman'

    # --- Case caption ---
    caption_para = doc.add_paragraph()
    caption_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption_para.paragraph_format.space_after = Pt(6)
    cap_run = caption_para.add_run('Case No. 2025-CV-04821')
    cap_run.font.size = Pt(12)
    cap_run.font.name = 'Times New Roman'

    parties_para = doc.add_paragraph()
    parties_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    parties_para.paragraph_format.space_after = Pt(12)
    p_run = parties_para.add_run('RIVERDALE MANUFACTURING, INC., Plaintiff,\nv.\nPACIFIC COAST SUPPLY CO., Defendant.')
    p_run.font.size = Pt(12)
    p_run.font.name = 'Times New Roman'

    # --- Separator ---
    sep_para = doc.add_paragraph()
    sep_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sep_para.paragraph_format.space_after = Pt(12)
    sep_run = sep_para.add_run('_' * 50)
    sep_run.font.size = Pt(12)
    sep_run.font.name = 'Times New Roman'

    # --- Body paragraphs ---
    body_style_kwargs = {'size': Pt(12), 'name': 'Times New Roman'}

    intro = doc.add_paragraph()
    intro.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    intro.paragraph_format.space_after = Pt(6)
    intro.paragraph_format.first_line_indent = Inches(0.5)
    r = intro.add_run(
        'COMES NOW the Plaintiff, Riverdale Manufacturing, Inc., by and through '
        'its attorneys of record, Henderson & Whitfield LLP, and hereby moves this '
        'Honorable Court for an Order compelling the Defendant, Pacific Coast Supply Co., '
        'to produce documents and respond to interrogatories previously served upon '
        'Defendant on January 15, 2025, and in support thereof states as follows:'
    )
    r.font.size = body_style_kwargs['size']
    r.font.name = body_style_kwargs['name']

    # Section I
    sec1_head = doc.add_paragraph()
    sec1_head.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sec1_head.paragraph_format.space_before = Pt(12)
    sec1_head.paragraph_format.space_after = Pt(6)
    s1r = sec1_head.add_run('I. BACKGROUND')
    s1r.bold = True
    s1r.font.size = Pt(12)
    s1r.font.name = 'Times New Roman'

    bg1 = doc.add_paragraph()
    bg1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    bg1.paragraph_format.space_after = Pt(6)
    bg1.paragraph_format.first_line_indent = Inches(0.5)
    r = bg1.add_run(
        '1. On September 3, 2024, Plaintiff filed the instant action alleging breach '
        'of contract and fraudulent misrepresentation arising from a supply agreement '
        'entered into on March 12, 2023, under which Defendant agreed to deliver '
        'industrial-grade polymer components to Plaintiff\'s Riverside facility.'
    )
    r.font.size = body_style_kwargs['size']
    r.font.name = body_style_kwargs['name']

    bg2 = doc.add_paragraph()
    bg2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    bg2.paragraph_format.space_after = Pt(6)
    bg2.paragraph_format.first_line_indent = Inches(0.5)
    r = bg2.add_run(
        '2. On January 15, 2025, Plaintiff served upon Defendant its First Set of '
        'Interrogatories (Nos. 1 through 18) and First Request for Production of '
        'Documents (Nos. 1 through 24). Defendant\'s responses were due on or before '
        'February 14, 2025, pursuant to Rule 33(b)(2) and Rule 34(b)(2)(A).'
    )
    r.font.size = body_style_kwargs['size']
    r.font.name = body_style_kwargs['name']

    bg3 = doc.add_paragraph()
    bg3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    bg3.paragraph_format.space_after = Pt(6)
    bg3.paragraph_format.first_line_indent = Inches(0.5)
    r = bg3.add_run(
        '3. As of the date of this filing, Defendant has failed to provide any responses '
        'to the aforementioned discovery requests, despite Plaintiff\'s counsel sending '
        'two written reminders dated February 20, 2025, and March 5, 2025, respectively.'
    )
    r.font.size = body_style_kwargs['size']
    r.font.name = body_style_kwargs['name']

    # Section II
    sec2_head = doc.add_paragraph()
    sec2_head.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sec2_head.paragraph_format.space_before = Pt(12)
    sec2_head.paragraph_format.space_after = Pt(6)
    s2r = sec2_head.add_run('II. ARGUMENT')
    s2r.bold = True
    s2r.font.size = Pt(12)
    s2r.font.name = 'Times New Roman'

    arg1 = doc.add_paragraph()
    arg1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    arg1.paragraph_format.space_after = Pt(6)
    arg1.paragraph_format.first_line_indent = Inches(0.5)
    r = arg1.add_run(
        '4. Under Federal Rule of Civil Procedure 37(a)(3)(B), a party seeking '
        'discovery may move for an order compelling an answer or production when an '
        'opposing party fails to respond to interrogatories or requests for production. '
        'The moving party must certify that it has in good faith conferred or attempted '
        'to confer with the non-responding party before filing such motion.'
    )
    r.font.size = body_style_kwargs['size']
    r.font.name = body_style_kwargs['name']

    arg2 = doc.add_paragraph()
    arg2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    arg2.paragraph_format.space_after = Pt(6)
    arg2.paragraph_format.first_line_indent = Inches(0.5)
    r = arg2.add_run(
        '5. Plaintiff has satisfied the meet-and-confer requirement. Counsel for '
        'Plaintiff contacted Defendant\'s counsel by telephone on February 28, 2025, '
        'and again on March 10, 2025. During both conversations, Defendant\'s counsel '
        'acknowledged the outstanding discovery but failed to provide a date by which '
        'responses would be served. See Exhibit A (Declaration of Attorney Rachel Nguyen).'
    )
    r.font.size = body_style_kwargs['size']
    r.font.name = body_style_kwargs['name']

    # Section III
    sec3_head = doc.add_paragraph()
    sec3_head.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sec3_head.paragraph_format.space_before = Pt(12)
    sec3_head.paragraph_format.space_after = Pt(6)
    s3r = sec3_head.add_run('III. CONCLUSION')
    s3r.bold = True
    s3r.font.size = Pt(12)
    s3r.font.name = 'Times New Roman'

    concl = doc.add_paragraph()
    concl.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    concl.paragraph_format.space_after = Pt(6)
    concl.paragraph_format.first_line_indent = Inches(0.5)
    r = concl.add_run(
        'WHEREFORE, Plaintiff Riverdale Manufacturing, Inc. respectfully requests that '
        'this Court enter an Order compelling Defendant Pacific Coast Supply Co. to respond '
        'fully to Plaintiff\'s First Set of Interrogatories and First Request for Production '
        'of Documents within fourteen (14) days of the date of such Order, and for such '
        'other and further relief as this Court deems just and proper.'
    )
    r.font.size = body_style_kwargs['size']
    r.font.name = body_style_kwargs['name']

    # Signature block
    sig_para = doc.add_paragraph()
    sig_para.paragraph_format.space_before = Pt(24)
    sig_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    r = sig_para.add_run(
        'Respectfully submitted,\n\n'
        'HENDERSON & WHITFIELD LLP\n\n'
        '_______________________________\n'
        'Rachel Nguyen, Esq.\n'
        'Bar No. 287415\n'
        '1200 Commerce Tower, Suite 1500\n'
        'Los Angeles, CA 90017\n'
        'Tel: (213) 555-0142\n'
        'rnguyen@hendersonwhitfield.com\n\n'
        'Attorneys for Plaintiff\n'
        'Riverdale Manufacturing, Inc.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
