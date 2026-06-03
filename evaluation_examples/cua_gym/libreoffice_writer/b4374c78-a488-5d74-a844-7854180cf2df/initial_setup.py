"""
Initial Setup: Business letter with mixed font sizes (12pt body, 16pt headings, 9pt footnotes)
Task ID: writer_frd_033
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_033'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_paragraph(doc, text, level=1):
    """Add a heading-style paragraph at 16pt."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Liberation Sans'
    return para


def add_body_paragraph(doc, text, alignment=None):
    """Add a body paragraph at 12pt."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    if alignment:
        para.paragraph_format.alignment = alignment
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Liberation Serif'
    return para


def add_footnote_paragraph(doc, text):
    """Add a footnote-style paragraph at 9pt."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Liberation Serif'
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return para


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # === Company Letterhead ===
    add_heading_paragraph(doc, 'Meridian Consulting Group')

    add_body_paragraph(
        doc,
        '1450 Technology Drive, Suite 300\nSan Francisco, CA 94107\nTel: (415) 555-0192 | Fax: (415) 555-0193'
    )

    add_body_paragraph(doc, 'March 28, 2026')

    add_body_paragraph(doc, '')

    # === Recipient ===
    add_body_paragraph(doc, 'Ms. Elena Vasquez\nDirector of Operations\nPacific Rim Industries\n8800 Harbor Boulevard\nLong Beach, CA 90802')

    add_body_paragraph(doc, 'Dear Ms. Vasquez,')

    # === Subject Heading ===
    add_heading_paragraph(doc, 'Re: Q1 2026 Operational Efficiency Assessment')

    # === Body Paragraphs ===
    add_body_paragraph(
        doc,
        'Thank you for engaging Meridian Consulting Group to conduct the operational '
        'efficiency assessment for Pacific Rim Industries. We are pleased to present our '
        'preliminary findings following the on-site review completed on March 15, 2026.'
    )

    add_body_paragraph(
        doc,
        'Our analysis covered three primary areas: warehouse logistics, procurement '
        'workflows, and interdepartmental communication channels. Each area was evaluated '
        'against industry benchmarks established by the National Operations Council.'
    )

    add_heading_paragraph(doc, 'Key Findings')

    add_body_paragraph(
        doc,
        '1. Warehouse Throughput: Current processing capacity stands at approximately '
        '2,340 units per shift, which is 18% below the benchmark of 2,850 units. The '
        'primary bottleneck appears to be the manual sorting stage between receiving and '
        'shelving operations.'
    )

    add_body_paragraph(
        doc,
        '2. Procurement Cycle Time: The average procurement cycle from requisition to '
        'delivery is 14.3 business days. Industry best practice suggests a target of 9 to '
        '10 business days. We identified redundant approval steps in the vendor selection '
        'process that account for approximately 3.2 days of delay.'
    )

    add_body_paragraph(
        doc,
        '3. Communication Efficiency: Cross-departmental response times average 6.8 hours '
        'for priority requests. Implementing the proposed unified ticketing system could '
        'reduce this to under 2 hours based on comparable deployments at similar organizations.'
    )

    add_heading_paragraph(doc, 'Recommendations')

    add_body_paragraph(
        doc,
        'We recommend a phased implementation approach beginning with the procurement '
        'workflow optimization, as it offers the highest return on investment with an '
        'estimated annual savings of $287,000. The warehouse automation upgrade should '
        'follow in Phase 2, with projected savings of $415,000 annually once fully deployed.'
    )

    add_body_paragraph(
        doc,
        'Our team is available to discuss these findings in detail at your convenience. '
        'We suggest scheduling a follow-up meeting during the week of April 7 to review '
        'the complete report and discuss implementation timelines.'
    )

    add_body_paragraph(doc, 'Sincerely,')

    add_body_paragraph(doc, '')

    add_body_paragraph(doc, 'Robert Chen\nSenior Partner\nMeridian Consulting Group')

    # === Footnotes at bottom ===
    add_footnote_paragraph(
        doc,
        '1. All figures cited in this letter are based on data collected during the '
        'period of February 10 through March 15, 2026.'
    )

    add_footnote_paragraph(
        doc,
        '2. Industry benchmarks sourced from the National Operations Council Annual '
        'Report, 2025 Edition (NOC-AR-2025).'
    )

    add_footnote_paragraph(
        doc,
        '3. Projected savings estimates assume full implementation within 12 months '
        'and are subject to revision upon detailed cost analysis.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
