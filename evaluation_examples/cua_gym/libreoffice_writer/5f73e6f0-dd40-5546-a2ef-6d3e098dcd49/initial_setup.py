"""
Initial Setup: Insert page number in header with Roman numerals
Task ID: writer_tm_085
Domain: libreoffice_writer

Creates a 5-page executive summary document with a left-aligned header
containing company logo text. No page numbers present.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_085'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Header: left-aligned company logo text, NO page numbers ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = hp.add_run("Meridian Corp")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # --- Page 1: Title and Introduction ---
    title = doc.add_heading("Executive Summary", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = subtitle.add_run("Q4 2025 Strategic Review")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    doc.add_paragraph(
        "This executive summary presents a comprehensive overview of Meridian Corp's "
        "strategic initiatives, financial performance, and operational milestones achieved "
        "during the fourth quarter of 2025. The document is intended for the Board of "
        "Directors and senior leadership team to facilitate informed decision-making for "
        "the upcoming fiscal year."
    )

    doc.add_paragraph(
        "Meridian Corp has demonstrated strong growth across all major business segments, "
        "with consolidated revenue reaching $487.3 million, representing a 12.4% increase "
        "compared to the same period last year. Our strategic investments in digital "
        "transformation, talent acquisition, and market expansion have positioned the "
        "company favorably for continued growth in 2026."
    )

    # --- Page 2: Financial Performance ---
    doc.add_page_break()
    doc.add_heading("Financial Performance", level=1)

    doc.add_paragraph(
        "The financial results for Q4 2025 exceeded analyst expectations across all "
        "key metrics. Revenue growth was primarily driven by the Technology Solutions "
        "division, which recorded a 23% year-over-year increase, and the Professional "
        "Services segment, which expanded its client base by 18%."
    )

    doc.add_heading("Revenue Breakdown by Division", level=2)
    doc.add_paragraph(
        "Technology Solutions contributed $198.7 million in revenue, representing 40.8% "
        "of total consolidated revenue. The Professional Services division generated "
        "$142.5 million, while the Enterprise Consulting group brought in $89.4 million. "
        "The remaining $56.7 million came from licensing and maintenance agreements."
    )

    doc.add_paragraph(
        "Operating margins improved to 18.3%, up from 15.7% in Q4 2024. This improvement "
        "was attributed to efficiency gains from the company-wide automation initiative "
        "launched in early 2025, which reduced operational overhead by approximately "
        "$12.8 million on an annualized basis."
    )

    doc.add_heading("Cash Flow and Balance Sheet", level=2)
    doc.add_paragraph(
        "Free cash flow for the quarter totaled $67.2 million, bringing the full-year "
        "total to $245.8 million. The company maintains a strong balance sheet with "
        "$312.4 million in cash and short-term investments, and a debt-to-equity ratio "
        "of 0.42, well within the target range of 0.3 to 0.6."
    )

    # --- Page 3: Strategic Initiatives ---
    doc.add_page_break()
    doc.add_heading("Strategic Initiatives", level=1)

    doc.add_paragraph(
        "Several critical strategic initiatives were advanced during Q4 2025, aligning "
        "with the company's five-year growth plan approved by the Board in January 2024."
    )

    doc.add_heading("Digital Transformation Program", level=2)
    doc.add_paragraph(
        "The Aurora digital transformation program completed Phase 2 implementation, "
        "migrating 78% of legacy systems to the new cloud-native architecture. Customer "
        "onboarding time decreased by 34%, and system uptime improved to 99.97%. The "
        "remaining Phase 3 activities are scheduled for completion by June 2026."
    )

    doc.add_heading("Market Expansion", level=2)
    doc.add_paragraph(
        "The Asia-Pacific expansion strategy yielded significant results, with new offices "
        "established in Singapore and Sydney. The region generated $28.3 million in new "
        "contract value during the quarter, surpassing the annual target of $25 million. "
        "Key client acquisitions include Tanaka Industries, Oceanic Holdings, and Pacific "
        "Rim Financial Group."
    )

    doc.add_paragraph(
        "In the European market, the acquisition of Nordic Solutions AB was completed in "
        "November 2025, adding 145 employees and approximately $35 million in annual "
        "recurring revenue. Integration activities are progressing on schedule, with full "
        "operational integration expected by Q2 2026."
    )

    # --- Page 4: Operational Highlights ---
    doc.add_page_break()
    doc.add_heading("Operational Highlights", level=1)

    doc.add_heading("Talent and Workforce", level=2)
    doc.add_paragraph(
        "Total headcount reached 4,287 employees as of December 31, 2025, a net increase "
        "of 412 positions during the quarter. Employee retention improved to 91.2%, "
        "compared to the industry average of 84.5%. The company's investment in the "
        "Meridian Academy training program contributed to a 28% increase in internal "
        "promotions year-over-year."
    )

    doc.add_heading("Client Satisfaction", level=2)
    doc.add_paragraph(
        "Net Promoter Score improved to 72, up from 65 in the prior quarter. Client "
        "satisfaction surveys indicated particularly high marks in project delivery "
        "timeliness (4.6/5.0) and technical expertise (4.8/5.0). The number of active "
        "enterprise clients grew to 347, with an average contract value of $1.4 million."
    )

    doc.add_heading("Technology Infrastructure", level=2)
    doc.add_paragraph(
        "The company completed migration of its primary data center to a hybrid cloud "
        "environment, reducing infrastructure costs by 22% while improving disaster "
        "recovery capabilities. The new architecture supports automatic scaling during "
        "peak periods, handling up to 150% of baseline traffic without performance "
        "degradation."
    )

    # --- Page 5: Outlook and Recommendations ---
    doc.add_page_break()
    doc.add_heading("Outlook and Recommendations", level=1)

    doc.add_paragraph(
        "Based on current market conditions and the strong foundation established in 2025, "
        "management projects consolidated revenue of $2.1 billion for fiscal year 2026, "
        "representing approximately 14% growth over the prior year. Operating margins are "
        "expected to reach 19.5% as efficiency initiatives continue to mature."
    )

    doc.add_heading("Key Priorities for 2026", level=2)

    doc.add_paragraph("Complete Phase 3 of the Aurora digital transformation program", style="List Bullet")
    doc.add_paragraph("Expand Asia-Pacific operations to include Tokyo and Mumbai offices", style="List Bullet")
    doc.add_paragraph("Launch the Meridian AI Suite product line targeting enterprise clients", style="List Bullet")
    doc.add_paragraph("Achieve ISO 27001 certification across all global operations", style="List Bullet")
    doc.add_paragraph("Increase R&D investment to 8% of revenue from current 6.5%", style="List Bullet")

    doc.add_paragraph("")

    doc.add_paragraph(
        "The Board is requested to review and approve the proposed strategic budget "
        "allocation of $315 million for the initiatives outlined above. Detailed financial "
        "models and risk assessments for each initiative are provided in the appendices "
        "to this summary."
    )

    doc.add_paragraph(
        "Respectfully submitted by the Office of the Chief Executive Officer, "
        "Meridian Corp, January 15, 2026."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
