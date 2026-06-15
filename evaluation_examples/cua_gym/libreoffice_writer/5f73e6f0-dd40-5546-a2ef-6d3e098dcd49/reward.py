"""
Reward Script: Insert right-aligned uppercase Roman numeral page number in header
Task ID: writer_tm_085
Domain: libreoffice_writer
Scoring:
  - Precondition gate: Original "Meridian Corp" text must be preserved (0.0 if missing)
  - Component 1 (0.40): PAGE field code exists in header
  - Component 2 (0.40): PAGE field uses uppercase Roman numeral format (\* ROMAN)
  - Component 3 (0.20): Right-aligned tab stop exists for positioning
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_085'


def persist_app_state(domain: str):
    """Best-effort save via Ctrl+S for GUI apps."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # We need to check the header of the first section
    try:
        section = doc.sections[0]
        header = section.header
        if not header.paragraphs:
            print("FAIL: No paragraphs found in header")
            print("REWARD: 0.0")
            return 0.0
    except Exception as e:
        print(f"CRITICAL: Cannot access header: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Precondition gate: Original "Meridian Corp" text must be preserved
    # This is a pre-existing property — NOT a scoring component, but a gate
    try:
        header_full_text = ' '.join(p.text for p in header.paragraphs)
        if 'Meridian Corp' not in header_full_text:
            print(f"PRECONDITION FAIL: 'Meridian Corp' not found in header — original content destroyed")
            print(f"  Header text: \"{header_full_text}\"")
            print("REWARD: 0.0")
            return 0.0
        else:
            print(f"PRECONDITION OK: 'Meridian Corp' preserved in header")
    except Exception as e:
        print(f"PRECONDITION ERROR: {e}")

    # Component 1: PAGE field code exists in header (0.40 points)
    # The header must contain a field code with instrText containing "PAGE"
    try:
        has_page_field = False
        for para in header.paragraphs:
            instr_elements = para._element.findall('.//w:instrText', ns)
            for instr in instr_elements:
                if instr.text and 'PAGE' in instr.text.upper():
                    has_page_field = True
                    break
            if has_page_field:
                break

        if has_page_field:
            print(f"PASS: Component 1 — PAGE field code found in header (0.40 pts)")
            total_score += 0.40
        else:
            print(f"FAIL: Component 1 — No PAGE field code found in header")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: PAGE field uses uppercase Roman numeral format (0.40 points)
    # The instrText should contain \* ROMAN (uppercase Roman) not \* roman (lowercase)
    try:
        has_roman_format = False
        for para in header.paragraphs:
            instr_elements = para._element.findall('.//w:instrText', ns)
            for instr in instr_elements:
                if instr.text and 'PAGE' in instr.text.upper():
                    instr_text = instr.text
                    # Check for \* ROMAN (uppercase Roman numerals)
                    # Accept both "ROMAN" and "Roman" as valid uppercase format
                    # but NOT "roman" (which would be lowercase i, ii, iii)
                    if '\\* ROMAN' in instr_text or '\\* Roman' in instr_text:
                        has_roman_format = True
                    elif '\\*ROMAN' in instr_text or '\\*Roman' in instr_text:
                        has_roman_format = True
                    print(f"  instrText content: \"{instr_text}\"")
                    break

        if has_roman_format:
            print(f"PASS: Component 2 — Uppercase Roman numeral format found (0.40 pts)")
            total_score += 0.40
        else:
            print(f"FAIL: Component 2 — No uppercase Roman format switch found in PAGE field")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Right-aligned positioning (0.20 points)
    # There should be a right-aligned tab stop in the header paragraph,
    # OR the paragraph has a right-alignment element for the page number portion
    try:
        has_right_alignment = False
        for para in header.paragraphs:
            # Check for right-aligned tab stops
            for ts in para.paragraph_format.tab_stops:
                from docx.enum.text import WD_TAB_ALIGNMENT
                if ts.alignment == WD_TAB_ALIGNMENT.RIGHT:
                    has_right_alignment = True
                    print(f"  Found right tab stop at position {ts.position}")
                    break

            # Also check XML for tab stops with val="right"
            if not has_right_alignment:
                tabs = para._element.findall('.//w:tabs/w:tab', ns)
                for tab in tabs:
                    val = tab.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if val == 'right':
                        has_right_alignment = True
                        print(f"  Found right tab in XML")
                        break

            if has_right_alignment:
                break

        if has_right_alignment:
            print(f"PASS: Component 3 — Right-aligned positioning found (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — No right-aligned tab stop or alignment found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
