"""
Initial Setup: Resume file for Michael Torres on the Desktop
Task ID: writer_creative_032
Domain: libreoffice_writer + OS file management
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_032'
OUTPUT = f'{WORKDIR}/Desktop/resume_old.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    desktop_dir = f'{WORKDIR}/Desktop'
    os.makedirs(desktop_dir, exist_ok=True)

    # Remove any pre-existing renamed version to ensure clean initial state
    golden_file = f'{desktop_dir}/Torres_Michael_Resume_2026.docx'
    if os.path.exists(golden_file):
        os.remove(golden_file)

    doc = Document()

    # ---- Name / Header ----
    name_para = doc.add_paragraph()
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_run = name_para.add_run('Michael Torres')
    name_run.bold = True
    name_run.font.size = Pt(18)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_para.add_run('San Francisco, CA  |  (415) 555-0192  |  m.torres@email.com  |  linkedin.com/in/michaeltorres')

    doc.add_paragraph()  # spacer

    # ---- Professional Summary ----
    summary_heading = doc.add_paragraph()
    sh_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
    sh_run.bold = True
    sh_run.font.size = Pt(12)

    doc.add_paragraph(
        'Results-driven software engineer with 7+ years of experience designing and delivering '
        'scalable backend systems and cloud-native applications. Proven track record of leading '
        'cross-functional teams to deliver complex projects on time. Strong expertise in Python, '
        'distributed systems, and AWS infrastructure.'
    )

    doc.add_paragraph()  # spacer

    # ---- Work Experience ----
    exp_heading = doc.add_paragraph()
    eh_run = exp_heading.add_run('WORK EXPERIENCE')
    eh_run.bold = True
    eh_run.font.size = Pt(12)

    # Job 1
    j1_para = doc.add_paragraph()
    j1_title = j1_para.add_run('Senior Software Engineer — Cloudify Inc., San Francisco, CA')
    j1_title.bold = True
    doc.add_paragraph('March 2021 – Present')
    doc.add_paragraph(
        '• Led architecture redesign of the core data pipeline, reducing processing latency by 40%.',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Managed a team of 5 engineers across two time zones to deliver a microservices migration on schedule.',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Implemented CI/CD pipelines using GitHub Actions and AWS CodeDeploy, cutting release cycles from 2 weeks to 3 days.',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Designed RESTful APIs consumed by 12 downstream services handling over 2 million daily requests.',
        style='List Bullet'
    )

    doc.add_paragraph()  # spacer

    # Job 2
    j2_para = doc.add_paragraph()
    j2_title = j2_para.add_run('Software Engineer — NexaTech Solutions, Austin, TX')
    j2_title.bold = True
    doc.add_paragraph('June 2018 – February 2021')
    doc.add_paragraph(
        '• Developed customer-facing features for a SaaS analytics platform serving 500+ enterprise clients.',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Optimized PostgreSQL query performance, improving dashboard load time by 35%.',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Collaborated with product and design teams using Agile/Scrum methodology.',
        style='List Bullet'
    )

    doc.add_paragraph()  # spacer

    # Job 3
    j3_para = doc.add_paragraph()
    j3_title = j3_para.add_run('Junior Developer — BrightPath Digital, Portland, OR')
    j3_title.bold = True
    doc.add_paragraph('July 2016 – May 2018')
    doc.add_paragraph(
        '• Built and maintained internal tools using Django and React for HR and finance departments.',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Wrote unit and integration tests achieving 85% code coverage on critical modules.',
        style='List Bullet'
    )

    doc.add_paragraph()  # spacer

    # ---- Education ----
    edu_heading = doc.add_paragraph()
    eduh_run = edu_heading.add_run('EDUCATION')
    eduh_run.bold = True
    eduh_run.font.size = Pt(12)

    doc.add_paragraph('B.S. Computer Science — University of California, Berkeley, CA  |  Graduated May 2016')
    doc.add_paragraph('GPA: 3.7/4.0  |  Dean\'s List (4 semesters)  |  Senior Thesis: Distributed Consensus Algorithms')

    doc.add_paragraph()  # spacer

    # ---- Skills ----
    skills_heading = doc.add_paragraph()
    skh_run = skills_heading.add_run('TECHNICAL SKILLS')
    skh_run.bold = True
    skh_run.font.size = Pt(12)

    doc.add_paragraph('Languages:    Python, Go, JavaScript, TypeScript, SQL')
    doc.add_paragraph('Frameworks:   Django, FastAPI, Node.js, React')
    doc.add_paragraph('Cloud & DevOps: AWS (EC2, S3, Lambda, RDS), Docker, Kubernetes, Terraform')
    doc.add_paragraph('Databases:    PostgreSQL, Redis, MongoDB, Elasticsearch')
    doc.add_paragraph('Tools:        Git, Jira, Confluence, DataDog, PagerDuty')

    doc.add_paragraph()  # spacer

    # ---- Certifications ----
    cert_heading = doc.add_paragraph()
    certh_run = cert_heading.add_run('CERTIFICATIONS')
    certh_run.bold = True
    certh_run.font.size = Pt(12)

    doc.add_paragraph('• AWS Certified Solutions Architect – Associate (2023)', style='List Bullet')
    doc.add_paragraph('• Certified Kubernetes Administrator (CKA) (2022)', style='List Bullet')

    doc.add_paragraph()  # spacer

    # ---- Last line: last updated ----
    last_para = doc.add_paragraph('Last updated: January 2025')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
