"""
Initial Setup: Configure tracked changes display settings
Task ID: writer_lec_078
Domain: libreoffice_writer

Creates a Writer document with tracked changes using default display settings.
The document contains realistic content with both insertions and deletions tracked.
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_078'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # --- Title ---
    title = doc.add_heading('Quarterly Marketing Strategy Review', level=0)

    # --- Introduction ---
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This document outlines the marketing strategy for Q2 2025, '
        'incorporating feedback from the leadership team meeting held on March 12, 2025. '
        'Key areas of focus include digital advertising budget reallocation, '
        'social media engagement improvements, and partnership development with '
        'regional influencers.'
    )

    # --- Budget Section ---
    doc.add_heading('Budget Allocation', level=1)
    doc.add_paragraph(
        'The total marketing budget for Q2 2025 has been set at $425,000, '
        'representing a 15% increase from Q1. The breakdown is as follows:'
    )

    # Budget table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Category', 'Q1 Actual ($)', 'Q2 Planned ($)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    budget_data = [
        ['Digital Advertising', '98,500', '125,000'],
        ['Social Media Campaigns', '45,200', '62,000'],
        ['Content Production', '72,800', '85,000'],
        ['Events & Sponsorships', '55,000', '78,000'],
        ['Analytics & Tools', '18,500', '25,000'],
    ]
    for r, row_data in enumerate(budget_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')  # spacer

    # --- Strategy Section ---
    doc.add_heading('Digital Strategy Update', level=1)
    doc.add_paragraph(
        'Based on Q1 performance metrics, we recommend shifting 20% of the '
        'display advertising budget toward programmatic channels. Our A/B testing '
        'showed a 34% higher conversion rate through targeted programmatic placements '
        'compared to traditional display networks.'
    )
    doc.add_paragraph(
        'The social media team will focus on three primary platforms: LinkedIn for '
        'B2B lead generation, Instagram for brand awareness, and TikTok for reaching '
        'younger demographics in the 18-34 age segment.'
    )

    # --- Team Section ---
    doc.add_heading('Team Assignments', level=1)
    doc.add_paragraph(
        'Project leads for Q2 initiatives:'
    )
    doc.add_paragraph('Sarah Chen - Digital Advertising & Programmatic', style='List Bullet')
    doc.add_paragraph('Marcus Johnson - Social Media Strategy', style='List Bullet')
    doc.add_paragraph('Elena Rodriguez - Content Production & Brand Voice', style='List Bullet')
    doc.add_paragraph('David Park - Analytics & Performance Tracking', style='List Bullet')
    doc.add_paragraph('Aisha Williams - Events & Partnership Development', style='List Bullet')

    # --- Timeline Section ---
    doc.add_heading('Implementation Timeline', level=1)
    doc.add_paragraph(
        'Phase 1 (April 1-15): Platform audit and baseline metrics collection. '
        'Phase 2 (April 16-30): Campaign launch for digital and social channels. '
        'Phase 3 (May 1-31): Mid-quarter review and optimization. '
        'Phase 4 (June 1-30): Final push and Q2 close-out reporting.'
    )

    # --- Conclusion ---
    doc.add_heading('Next Steps', level=1)
    doc.add_paragraph(
        'The marketing leadership team will reconvene on April 5, 2025 to finalize '
        'vendor contracts and approve creative assets. All team leads are expected to '
        'submit their detailed execution plans by March 28, 2025.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Ensure default tracked changes settings (reset to defaults if needed)
    # The registrymodifications.xcu should NOT have custom Insert/Delete display settings
    reg_file = '/home/user/.config/libreoffice/4/user/registrymodifications.xcu'
    if os.path.exists(reg_file):
        with open(reg_file, 'r') as f:
            content = f.read()
        # Remove any existing tracked changes display settings to ensure defaults
        import re
        content = re.sub(
            r'<item oor:path="/org\.openoffice\.Office\.Writer/Revision/TextDisplay/Insert">[^<]*<prop[^>]*>[^<]*<value>[^<]*</value>[^<]*</prop>[^<]*</item>\n?',
            '', content
        )
        content = re.sub(
            r'<item oor:path="/org\.openoffice\.Office\.Writer/Revision/TextDisplay/Delete">[^<]*<prop[^>]*>[^<]*<value>[^<]*</value>[^<]*</prop>[^<]*</item>\n?',
            '', content
        )
        with open(reg_file, 'w') as f:
            f.write(content)
        print('Cleaned any existing tracked changes display settings')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
