"""
Initial Setup: 8-page global economics survey document (no endnotes)
Task ID: writer_struct_050
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'global_economics_survey'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    return para


def add_paragraph(doc, text, first_line_indent=True):
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if first_line_indent:
        para.paragraph_format.first_line_indent = Inches(0.5)
    return para


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # -----------------------------------------------------------------------
    # PAGE 1 — Chapter 1: Global GDP Trends
    # -----------------------------------------------------------------------
    heading = doc.add_heading('Global Economic Survey 2024', level=0)
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph('An Analysis of Growth, Policy, and Emerging Markets')
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(12)

    add_paragraph(doc,
        'Chapter 1: Global GDP Trends',
        first_line_indent=False
    ).runs[0].font.bold = True

    # FIRST PARAGRAPH on page 1 — endnote reference goes here
    add_paragraph(doc,
        'Global economic growth moderated to approximately 3.1 percent in 2023, '
        'reflecting the cumulative effects of monetary policy tightening across '
        'advanced economies, lingering supply chain disruptions, and heightened '
        'geopolitical uncertainties. The International Monetary Fund projects a '
        'slight recovery to 3.2 percent in 2024, contingent on a soft landing in '
        'major economies and a gradual easing of inflationary pressures. Divergence '
        'among regional blocs remains pronounced, with emerging market economies '
        'maintaining stronger growth trajectories relative to their developed-world '
        'counterparts.'
    )

    add_paragraph(doc,
        'Advanced economies faced significant headwinds during 2023. The United States '
        'recorded GDP growth of 2.5 percent, outperforming many forecasts, driven by '
        'robust consumer spending, a resilient labor market, and sustained business '
        'investment in technology and infrastructure. In contrast, the Euro Area '
        'grew by just 0.5 percent, constrained by elevated energy prices, weak '
        'manufacturing output in Germany, and fiscal consolidation pressures across '
        'southern member states. The United Kingdom narrowly avoided a technical '
        'recession, posting growth of 0.1 percent amid persistent inflation and '
        'declining real household incomes.'
    )

    add_paragraph(doc,
        'Asia-Pacific economies continued to demonstrate relative resilience. China\'s '
        'post-pandemic recovery delivered GDP growth of 5.2 percent, meeting official '
        'targets but underscoring structural challenges in the property sector and '
        'subdued domestic consumer confidence. India emerged as the fastest-growing '
        'major economy at 6.8 percent, supported by strong public investment, '
        'expanding digital infrastructure, and a booming services export sector. '
        'Japan achieved modest growth of 1.9 percent, benefiting from a weaker yen '
        'that boosted export revenues and tourism receipts.'
    )

    add_paragraph(doc,
        'Trade volumes contracted marginally in 2023 for the first time since the '
        '2020 pandemic year. The World Trade Organization attributed the decline to '
        'weaker goods demand in Europe and North America, ongoing reshoring and '
        'friend-shoring of supply chains, and the fragmenting of global trade networks '
        'along geopolitical fault lines. Services trade, however, remained buoyant, '
        'with international tourism and cross-border digital commerce registering '
        'double-digit growth in volume terms.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # PAGE 2 — Chapter 2: Fiscal Policy
    # -----------------------------------------------------------------------
    add_paragraph(doc,
        'Chapter 2: Fiscal Policy Developments',
        first_line_indent=False
    ).runs[0].font.bold = True

    # FIRST PARAGRAPH on page 2
    add_paragraph(doc,
        'Fiscal policy in 2023 was characterized by a gradual withdrawal of the '
        'extraordinary support measures introduced during the COVID-19 pandemic. '
        'Most advanced economies reduced their primary deficits, driven by '
        'expenditure restraint, improved tax revenues from wage growth, and the '
        'rolling off of emergency spending programs. Nevertheless, government debt '
        'ratios remained elevated in many countries, raising concerns about long-term '
        'debt sustainability and the adequacy of fiscal buffers ahead of potential '
        'future shocks.'
    )

    # SECOND PARAGRAPH on page 2 — endnote reference goes here
    add_paragraph(doc,
        'The composition of public expenditure shifted significantly across the '
        'OECD over the past year. Capital investment in green infrastructure, digital '
        'transformation, and defense modernization absorbed a growing share of '
        'national budgets, reflecting both security imperatives and the demands of '
        'the clean energy transition. Social protection spending, while declining as '
        'a share of GDP from pandemic peaks, remained above pre-2020 levels in most '
        'jurisdictions, providing ongoing support for lower-income households '
        'disproportionately affected by the cost-of-living crisis.'
    )

    add_paragraph(doc,
        'Tax policy reforms gained momentum in 2023, most notably with the phased '
        'implementation of the OECD/G20 global minimum corporate tax framework. '
        'Over 130 countries committed to the 15 percent minimum effective tax rate '
        'for large multinational corporations, with the first qualified domestic '
        'minimum top-up taxes entering into force at the start of 2024. Early '
        'revenue estimates suggest the reform could generate an additional $150 '
        'billion annually in global corporate tax receipts, though actual yields '
        'will depend heavily on implementation consistency and the extent of profit '
        'shifting mitigation.'
    )

    add_paragraph(doc,
        'Debt management challenges intensified as refinancing costs rose sharply '
        'across the maturity spectrum. Governments that had locked in low rates '
        'during the era of near-zero interest rates now face markedly higher rollover '
        'costs as legacy debt matures. In several highly indebted economies, the '
        'interest-to-revenue ratio exceeded 15 percent in 2023, constraining the '
        'fiscal space available for growth-enhancing public investments. Multilateral '
        'creditors and sovereign rating agencies intensified scrutiny of medium-term '
        'fiscal consolidation plans in response to these dynamics.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # PAGE 3 — Chapter 3: Monetary Policy
    # -----------------------------------------------------------------------
    add_paragraph(doc,
        'Chapter 3: Monetary Policy and Central Banking',
        first_line_indent=False
    ).runs[0].font.bold = True

    # FIRST PARAGRAPH on page 3 — endnote reference goes here
    add_paragraph(doc,
        'Central banks in major advanced economies entered 2023 in the midst of the '
        'most aggressive tightening cycle in four decades. The US Federal Reserve '
        'raised the federal funds rate to a range of 5.25 to 5.50 percent by July '
        '2023, the highest level since 2001. The European Central Bank lifted its '
        'deposit facility rate to 4.0 percent by September, while the Bank of England '
        'reached 5.25 percent. These coordinated tightening actions were effective in '
        'reducing headline inflation from its 2022 peaks, but core inflation proved '
        'stickier than anticipated, particularly in services sectors with persistent '
        'wage growth.'
    )

    add_paragraph(doc,
        'Transmission of monetary policy to the real economy proceeded with the '
        'familiar lags but also exhibited some atypical features in the current cycle. '
        'The housing market in the United States showed surprising resilience despite '
        'mortgage rates exceeding 7 percent, as existing homeowners locked into '
        'low-rate mortgages were reluctant to sell, creating a supply constraint that '
        'supported prices. Corporate borrowing costs rose sharply, yet investment-grade '
        'issuers maintained access to capital markets. Stress was most acute in '
        'commercial real estate, where falling valuations and maturity walls created '
        'concentrated risks within regional and community banking systems.'
    )

    add_paragraph(doc,
        'Emerging market central banks navigated a delicate balancing act throughout '
        '2023. Those that had tightened preemptively in 2021 and 2022, notably Brazil '
        'and Chile, began cautious easing cycles as domestic inflation declined while '
        'remaining alert to spillovers from Federal Reserve policy decisions. Asian '
        'emerging markets with current account surpluses and lower inflation were able '
        'to maintain accommodative stances for longer, supporting domestic demand. '
        'Frontier market economies with dollar-denominated debt faced the most acute '
        'stress, as higher US rates and dollar strength increased debt service burdens '
        'and widened sovereign spreads.'
    )

    add_paragraph(doc,
        'Central bank balance sheet normalization proceeded in parallel with rate '
        'increases. The Federal Reserve\'s quantitative tightening program reduced '
        'its holdings from a peak of approximately $9 trillion to under $7.5 trillion '
        'by the end of 2023 without significant market disruption. The European '
        'Central Bank concluded its Asset Purchase Programme reinvestments by mid-year. '
        'Academic debate intensified over the appropriate long-run size of central '
        'bank balance sheets and the feasibility of returning to pre-2008 frameworks '
        'given the structural changes in monetary transmission mechanisms.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # PAGE 4 — Chapter 4: Emerging Markets
    # -----------------------------------------------------------------------
    add_paragraph(doc,
        'Chapter 4: Emerging Market Dynamics',
        first_line_indent=False
    ).runs[0].font.bold = True

    # FIRST PARAGRAPH on page 4 — endnote reference goes here
    add_paragraph(doc,
        'Emerging market and developing economies as a group outpaced advanced '
        'economies in 2023, collectively recording GDP growth of approximately 4.3 '
        'percent. This aggregate figure, however, obscures substantial heterogeneity '
        'across regions and individual countries. Sub-Saharan Africa grew by 3.5 '
        'percent, held back by debt distress in several frontier markets, adverse '
        'weather events linked to El Niño, and commodity price volatility. Latin '
        'America expanded by just 2.3 percent as higher-for-longer interest rates, '
        'political uncertainty, and sluggish productivity growth weighed on activity.'
    )

    add_paragraph(doc,
        'Capital flows to emerging markets remained volatile throughout 2023. Net '
        'portfolio outflows occurred in the first and third quarters as global risk '
        'appetite fluctuated with US interest rate expectations. Foreign direct '
        'investment flows showed greater resilience, driven by near-shoring and '
        'friend-shoring trends that redirected manufacturing investment toward '
        'geopolitically aligned jurisdictions. Mexico, Vietnam, and India were among '
        'the principal beneficiaries of supply chain diversification, attracting '
        'record FDI inflows in electronics and semiconductor-adjacent manufacturing.'
    )

    add_paragraph(doc,
        'Currency markets in emerging economies experienced significant turbulence. '
        'Several currencies depreciated materially against the US dollar in periods '
        'of risk-off sentiment, necessitating central bank interventions. The '
        'Argentine peso lost over 50 percent of its value against the dollar during '
        'the year following the election of a new government that pursued a dramatic '
        'devaluation as part of a stabilization program. The Nigerian naira was '
        'devalued by approximately 40 percent after the removal of longstanding '
        'currency controls, while the Egyptian pound remained under pressure amid '
        'acute foreign exchange shortages.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # PAGE 5 — Chapter 5: Trade and Geopolitics
    # -----------------------------------------------------------------------
    add_paragraph(doc,
        'Chapter 5: Trade, Geopolitics, and Supply Chains',
        first_line_indent=False
    ).runs[0].font.bold = True

    add_paragraph(doc,
        'The geopolitical fracturing of global trade continued to reshape supply '
        'chain architectures in 2023. The United States expanded trade and technology '
        'restrictions targeting China, while the European Union advanced its '
        'Economic Security Strategy and Foreign Subsidies Regulation, signaling a '
        'more assertive posture on trade defense. China maintained reciprocal '
        'measures, including export controls on critical minerals essential to the '
        'clean energy transition. These developments amplified uncertainty for '
        'multinational firms dependent on cross-border value chains.'
    )

    add_paragraph(doc,
        'The ASEAN bloc emerged as a key beneficiary of shifting trade patterns. '
        'Vietnam, Indonesia, Thailand, and Malaysia attracted diversified manufacturing '
        'investments from electronics, apparel, and automotive sectors. Trade between '
        'ASEAN and both the United States and China grew in 2023, reflecting the '
        'region\'s strategic position as a neutral production hub capable of serving '
        'both great power blocs. Intra-ASEAN trade also deepened, supported by '
        'the implementation of the Regional Comprehensive Economic Partnership.'
    )

    add_paragraph(doc,
        'Food and energy security concerns prompted significant policy responses '
        'across the developing world. Following the disruption to Black Sea grain '
        'exports caused by the Russia-Ukraine conflict, over 40 countries implemented '
        'food export restrictions at some point in 2022 or 2023, exacerbating import '
        'price volatility for food-import-dependent nations. The energy transition '
        'introduced new forms of commodity dependency, as the demand for lithium, '
        'cobalt, copper, and rare earth elements surged in line with battery and '
        'renewable energy manufacturing capacity expansions.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # PAGE 6 — Chapter 6: Labor Markets
    # -----------------------------------------------------------------------
    add_paragraph(doc,
        'Chapter 6: Labor Markets and Human Capital',
        first_line_indent=False
    ).runs[0].font.bold = True

    add_paragraph(doc,
        'Labor market conditions remained unusually tight in advanced economies '
        'through much of 2023, defying expectations of rapid cooling in response '
        'to monetary tightening. Unemployment in the United States averaged just '
        '3.6 percent for the year, near historic lows, while job openings remained '
        'substantially above pre-pandemic levels. The Euro Area unemployment rate '
        'fell to 6.0 percent, the lowest since the creation of the single currency. '
        'These conditions supported wage growth that, while moderating, continued '
        'to run ahead of productivity gains in many sectors, sustaining inflationary '
        'pressures in labor-intensive services.'
    )

    add_paragraph(doc,
        'The structural transformation of labor markets accelerated in 2023 as '
        'generative artificial intelligence applications began to impact professional '
        'services, customer support, content creation, and software development. '
        'Early evidence suggests that AI adoption has been productivity-enhancing '
        'for skilled workers while reducing entry-level demand in certain occupational '
        'categories. Policy responses ranged from voluntary industry codes of conduct '
        'to binding regulatory frameworks, reflecting divergent national approaches '
        'to governing the labor market implications of AI deployment.'
    )

    add_paragraph(doc,
        'Demographic trends continue to constrain long-run growth potential across '
        'many advanced and some emerging economies. Japan, South Korea, Italy, and '
        'Germany face shrinking working-age populations that challenge pension '
        'sustainability and fiscal balance. Migration policy has emerged as a de '
        'facto labor market instrument, with Germany, Canada, and Australia expanding '
        'skilled worker immigration programs to address specific sectoral shortages. '
        'Meanwhile, developing countries with youthful demographic profiles, '
        'particularly in sub-Saharan Africa and South Asia, must create sufficient '
        'employment opportunities to absorb growing labor forces productively.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # PAGE 7 — Chapter 7: Financial Markets
    # -----------------------------------------------------------------------
    add_paragraph(doc,
        'Chapter 7: Financial Markets and Stability',
        first_line_indent=False
    ).runs[0].font.bold = True

    add_paragraph(doc,
        'Global equity markets delivered strong returns in 2023, belying earlier '
        'recession fears. The S&P 500 gained 24 percent for the year, led by the '
        '"Magnificent Seven" technology stocks whose valuations surged on enthusiasm '
        'for artificial intelligence applications. European equities rose modestly '
        'in local currency terms but underperformed on a dollar basis due to euro '
        'depreciation. Emerging market equities lagged, with the MSCI Emerging '
        'Markets Index posting a modest gain as Chinese equity underperformance '
        'offset strong returns in India, Taiwan, and South Korea.'
    )

    add_paragraph(doc,
        'Bond markets experienced exceptional volatility. The 10-year US Treasury '
        'yield briefly touched 5.0 percent in October 2023, a level not seen since '
        '2007, before retreating to close the year near 3.9 percent. This volatility '
        'reflected ongoing uncertainty about the Federal Reserve\'s terminal rate, '
        'the sustainability of high US deficits, and the potential shift in global '
        'reserve allocation patterns. Credit spreads remained surprisingly narrow '
        'throughout the year, with both investment-grade and high-yield issuers '
        'continuing to access markets at spreads inconsistent with the prevailing '
        'interest rate environment.'
    )

    add_paragraph(doc,
        'The global banking system demonstrated resilience in 2023, though two '
        'episodes revealed pockets of vulnerability. The failures of Silicon Valley '
        'Bank, Signature Bank, and Silvergate in the United States in March 2023 '
        'highlighted the duration risk embedded in banks that had invested heavily '
        'in long-dated fixed income securities during the low-rate era. The '
        'subsequent emergency rescue of Credit Suisse by UBS in Switzerland, '
        'orchestrated by Swiss regulators and the Swiss National Bank, underscored '
        'the continued too-big-to-fail concerns that persist despite the post-2008 '
        'regulatory overhaul. Both episodes triggered temporary contagion but were '
        'ultimately contained through swift official action.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # PAGE 8 — Chapter 8: Outlook and Conclusions
    # -----------------------------------------------------------------------
    add_paragraph(doc,
        'Chapter 8: Outlook and Policy Conclusions',
        first_line_indent=False
    ).runs[0].font.bold = True

    add_paragraph(doc,
        'The global economic outlook for 2024 and 2025 is shaped by several '
        'cross-cutting forces. On the positive side, the disinflation trend appears '
        'durable, creating space for central banks to begin gradual policy easing '
        'cycles. Labor markets remain supportive of household incomes. Fiscal '
        'stimulus embedded in landmark legislation such as the US Inflation Reduction '
        'Act and CHIPS and Science Act continues to crowd in private investment '
        'in clean energy and semiconductor manufacturing.'
    )

    add_paragraph(doc,
        'Downside risks, however, remain material. Geopolitical escalation — '
        'whether related to the Russia-Ukraine conflict, tensions over Taiwan, '
        'or Middle East instability — could generate adverse commodity price and '
        'financial market shocks. A harder-than-expected landing in the United '
        'States, triggered by lagged monetary policy effects, could dampen global '
        'demand and trade growth. Climate-related economic disruptions, increasingly '
        'apparent in insurance markets, agricultural yields, and infrastructure '
        'resilience, represent a growing medium-term risk factor that conventional '
        'macroeconomic models are only beginning to incorporate.'
    )

    add_paragraph(doc,
        'Policy priorities for the year ahead center on achieving a credible '
        'disinflation without triggering an unnecessary recession, rebuilding '
        'fiscal buffers to strengthen future resilience, advancing structural '
        'reforms to boost potential growth, and strengthening international '
        'coordination mechanisms to manage global commons challenges. The coming '
        'years will test the capacity of multilateral institutions and national '
        'policymakers to navigate an increasingly complex and fragmented global '
        'economic landscape.'
    )

    add_paragraph(doc,
        'The findings of this survey underscore the importance of continued '
        'investment in high-quality economic data collection, analysis, and '
        'international sharing frameworks. Policymakers, businesses, and '
        'civil society all require timely and reliable economic intelligence to '
        'navigate the uncertainties ahead. This survey, and the research it '
        'synthesizes, aims to contribute to that shared objective by providing '
        'a comprehensive and rigorous assessment of the state of the global '
        'economy at a critical juncture in its evolution.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
