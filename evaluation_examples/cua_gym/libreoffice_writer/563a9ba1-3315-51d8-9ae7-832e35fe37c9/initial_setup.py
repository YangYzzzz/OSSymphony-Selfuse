"""
Initial Setup: Sales Report with 8 pages; page 4 has 'Sales Visualization' heading and data table.
Task ID: writer_pd_037
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

# Install dependencies on VM
subprocess.run(['pip3', 'install', 'python-docx', 'Pillow'], capture_output=True)

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_037'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add an explicit page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)


def format_header_cell(cell, text):
    """Format a table header cell with bold white text on blue background."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, '2F5496')


def format_data_cell(cell, text, align_center=False):
    """Format a regular data cell."""
    cell.text = ''
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(str(text))
    run.font.size = Pt(10)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ============================================================
    # PAGE 1: Title Page
    # ============================================================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_heading('Quarterly Sales Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies Inc.')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    period = doc.add_paragraph()
    period.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = period.add_run('First Half 2026 (January - June)')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Prepared by: Sales Analytics Division\nDate: July 15, 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    add_page_break(doc)

    # ============================================================
    # PAGE 2: Executive Summary
    # ============================================================
    doc.add_heading('Executive Summary', level=1)

    doc.add_paragraph(
        'This report provides a comprehensive overview of sales performance for '
        'Meridian Technologies Inc. during the first half of 2026. The analysis covers '
        'monthly revenue trends, product category breakdowns, and regional performance metrics.'
    )

    doc.add_paragraph(
        'Total revenue for H1 2026 reached $2,847,500, representing a 12.3% increase '
        'compared to the same period in 2025. The strongest growth was observed in the '
        'Enterprise Solutions segment, which contributed 42% of total revenue.'
    )

    doc.add_heading('Key Highlights', level=2)
    highlights = [
        'Revenue exceeded quarterly targets by 8.7% in Q1 and 5.2% in Q2.',
        'Customer acquisition rate improved by 15% year-over-year.',
        'Average deal size increased from $18,400 to $22,750.',
        'Customer retention rate maintained at 94.2%.',
        'New market expansion into Southeast Asia contributed $187,000 in revenue.',
    ]
    for h in highlights:
        doc.add_paragraph(h, style='List Bullet')

    doc.add_paragraph(
        'The following sections present detailed monthly breakdowns, product category '
        'analysis, regional performance, and recommendations for the second half of the year.'
    )

    add_page_break(doc)

    # ============================================================
    # PAGE 3: Product Category Analysis
    # ============================================================
    doc.add_heading('Product Category Analysis', level=1)

    doc.add_paragraph(
        'Revenue distribution across product categories shows continued dominance of '
        'Enterprise Solutions, while the Cloud Services segment demonstrated the highest '
        'growth rate at 28.4% year-over-year.'
    )

    # Product category table
    cat_table = doc.add_table(rows=6, cols=4)
    cat_table.style = 'Table Grid'
    cat_headers = ['Category', 'H1 2026 Revenue', 'H1 2025 Revenue', 'YoY Growth']
    for i, h in enumerate(cat_headers):
        format_header_cell(cat_table.cell(0, i), h)

    cat_data = [
        ['Enterprise Solutions', '$1,195,950', '$1,042,300', '14.7%'],
        ['Cloud Services', '$569,500', '$443,400', '28.4%'],
        ['Professional Services', '$427,125', '$398,600', '7.2%'],
        ['Hardware & Licensing', '$398,650', '$382,100', '4.3%'],
        ['Support & Maintenance', '$256,275', '$269,800', '-5.0%'],
    ]
    for r, row_data in enumerate(cat_data, 1):
        for c, val in enumerate(row_data):
            format_data_cell(cat_table.cell(r, c), val, align_center=(c > 0))

    doc.add_paragraph()
    doc.add_paragraph(
        'Enterprise Solutions maintained its position as the primary revenue driver, '
        'accounting for 42% of total H1 revenue. The Cloud Services segment showed '
        'particularly strong momentum, driven by increased adoption of our SaaS platform '
        'among mid-market customers.'
    )

    doc.add_paragraph(
        'The decline in Support & Maintenance revenue is attributable to the migration '
        'of legacy support contracts to bundled Cloud Services packages, which is a '
        'strategically positive shift despite the apparent decrease.'
    )

    add_page_break(doc)

    # ============================================================
    # PAGE 4: Sales Visualization (with data table, NO chart)
    # ============================================================
    doc.add_heading('Sales Visualization', level=1)

    doc.add_paragraph(
        'The table below summarizes monthly sales figures for the first half of 2026. '
        'These figures represent total revenue across all product categories and regions.'
    )

    # Monthly sales data table
    sales_table = doc.add_table(rows=7, cols=2)
    sales_table.style = 'Table Grid'
    format_header_cell(sales_table.cell(0, 0), 'Month')
    format_header_cell(sales_table.cell(0, 1), 'Sales ($)')

    monthly_data = [
        ('January', '$412,300'),
        ('February', '$389,750'),
        ('March', '$467,200'),
        ('April', '$498,500'),
        ('May', '$521,850'),
        ('June', '$557,900'),
    ]
    for r, (month, sales) in enumerate(monthly_data, 1):
        format_data_cell(sales_table.cell(r, 0), month)
        format_data_cell(sales_table.cell(r, 1), sales, align_center=True)

    doc.add_paragraph()
    doc.add_paragraph(
        'Total H1 2026 Revenue: $2,847,500'
    ).runs[0].bold = True

    doc.add_paragraph(
        'The monthly progression shows a consistent upward trend from January through June, '
        'with the strongest month being June at $557,900. The average monthly revenue was '
        '$474,583, exceeding the target of $450,000.'
    )

    # Leave space below for chart insertion (the task asks the agent to add a chart here)
    doc.add_paragraph()
    doc.add_paragraph()

    add_page_break(doc)

    # ============================================================
    # PAGE 5: Regional Performance
    # ============================================================
    doc.add_heading('Regional Performance', level=1)

    doc.add_paragraph(
        'Sales performance varied significantly across regions, with North America '
        'continuing to lead while Asia-Pacific showed the strongest growth trajectory.'
    )

    reg_table = doc.add_table(rows=5, cols=3)
    reg_table.style = 'Table Grid'
    reg_headers = ['Region', 'Revenue', 'Share']
    for i, h in enumerate(reg_headers):
        format_header_cell(reg_table.cell(0, i), h)

    reg_data = [
        ['North America', '$1,423,750', '50.0%'],
        ['Europe', '$712,500', '25.0%'],
        ['Asia-Pacific', '$427,125', '15.0%'],
        ['Rest of World', '$284,125', '10.0%'],
    ]
    for r, row_data in enumerate(reg_data, 1):
        for c, val in enumerate(row_data):
            format_data_cell(reg_table.cell(r, c), val, align_center=(c > 0))

    doc.add_paragraph()
    doc.add_paragraph(
        'North America accounted for half of total revenue, driven by strong enterprise '
        'deals in the financial services and healthcare verticals. The Asia-Pacific region, '
        'while smaller in absolute terms, grew at 34% year-over-year, validating our '
        'expansion strategy in the region.'
    )

    add_page_break(doc)

    # ============================================================
    # PAGE 6: Top Accounts
    # ============================================================
    doc.add_heading('Top Accounts', level=1)

    doc.add_paragraph(
        'The following table lists the top ten accounts by revenue contribution during '
        'H1 2026. These accounts collectively represent 38% of total revenue.'
    )

    acct_table = doc.add_table(rows=11, cols=4)
    acct_table.style = 'Table Grid'
    acct_headers = ['Rank', 'Account Name', 'Industry', 'H1 Revenue']
    for i, h in enumerate(acct_headers):
        format_header_cell(acct_table.cell(0, i), h)

    acct_data = [
        ['1', 'Apex Financial Group', 'Financial Services', '$142,800'],
        ['2', 'NovaCare Health Systems', 'Healthcare', '$128,500'],
        ['3', 'Pinnacle Manufacturing', 'Manufacturing', '$119,300'],
        ['4', 'CrestPoint Energy', 'Energy & Utilities', '$112,700'],
        ['5', 'Silverline Retail Corp', 'Retail', '$104,200'],
        ['6', 'TerraVault Mining Ltd', 'Mining & Resources', '$98,500'],
        ['7', 'BlueStar Logistics', 'Transportation', '$91,800'],
        ['8', 'Quantum Dynamics R&D', 'Technology', '$87,400'],
        ['9', 'Harmony Education Trust', 'Education', '$79,600'],
        ['10', 'Pacific Rim Trading Co', 'International Trade', '$72,300'],
    ]
    for r, row_data in enumerate(acct_data, 1):
        for c, val in enumerate(row_data):
            format_data_cell(acct_table.cell(r, c), val, align_center=(c == 0 or c == 3))

    add_page_break(doc)

    # ============================================================
    # PAGE 7: Sales Team Performance
    # ============================================================
    doc.add_heading('Sales Team Performance', level=1)

    doc.add_paragraph(
        'Individual sales representative performance is tracked against quarterly quotas. '
        'The following summary highlights key performers and areas requiring attention.'
    )

    team_table = doc.add_table(rows=9, cols=4)
    team_table.style = 'Table Grid'
    team_headers = ['Representative', 'Territory', 'H1 Revenue', 'Quota Attainment']
    for i, h in enumerate(team_headers):
        format_header_cell(team_table.cell(0, i), h)

    team_data = [
        ['Sarah Chen', 'West Coast', '$387,200', '118%'],
        ['Marcus Johnson', 'East Coast', '$342,800', '104%'],
        ['Elena Rodriguez', 'Southeast', '$318,500', '97%'],
        ['David Kim', 'Midwest', '$298,400', '91%'],
        ['Priya Patel', 'Northeast', '$276,900', '109%'],
        ['James O\'Brien', 'Central', '$254,300', '95%'],
        ['Aisha Williams', 'Pacific NW', '$241,700', '112%'],
        ['Robert Tanaka', 'Southwest', '$228,100', '88%'],
    ]
    for r, row_data in enumerate(team_data, 1):
        for c, val in enumerate(row_data):
            format_data_cell(team_table.cell(r, c), val, align_center=(c >= 2))

    doc.add_paragraph()
    doc.add_paragraph(
        'Five of eight representatives exceeded or met their quota targets. Sarah Chen '
        'led the team with 118% quota attainment, driven by two major enterprise deals '
        'in the technology sector. Robert Tanaka\'s territory requires additional support '
        'and pipeline development initiatives.'
    )

    add_page_break(doc)

    # ============================================================
    # PAGE 8: Recommendations & Outlook
    # ============================================================
    doc.add_heading('Recommendations & Outlook', level=1)

    doc.add_heading('Strategic Recommendations', level=2)
    recs = [
        'Increase investment in Cloud Services marketing to capitalize on 28.4% growth momentum.',
        'Expand Asia-Pacific sales team by two additional representatives to support 34% regional growth.',
        'Implement cross-selling programs between Enterprise Solutions and Cloud Services.',
        'Develop targeted retention programs for Support & Maintenance customers transitioning to cloud.',
        'Launch partner channel program to extend reach in underperforming territories.',
    ]
    for r in recs:
        doc.add_paragraph(r, style='List Number')

    doc.add_heading('H2 2026 Outlook', level=2)
    doc.add_paragraph(
        'Based on current pipeline analysis and market conditions, we project H2 2026 '
        'revenue of $3,150,000 to $3,350,000, representing 10-17% growth over H2 2025. '
        'Key growth drivers include the anticipated launch of our next-generation cloud '
        'platform in September and expansion of the Southeast Asia partner ecosystem.'
    )

    doc.add_paragraph(
        'Risk factors include potential macroeconomic headwinds in European markets and '
        'increased competition in the mid-market segment. Mitigation strategies are being '
        'developed in conjunction with the product and marketing teams.'
    )

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
