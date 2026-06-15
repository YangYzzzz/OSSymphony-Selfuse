"""
Initial Setup: Create print_requirements.docx and artwork.jpg on the Desktop
Task ID: osworld_multi_apps_writer_to_gimp_014
Domain: multi_apps (libreoffice_writer + gimp)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont
import numpy as np

WORKDIR = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_writer_to_gimp_014'

REQUIREMENTS_DOC = f'{WORKDIR}/print_requirements.docx'
ARTWORK_PATH = f'{WORKDIR}/artwork.jpg'
PRINT_READY_PATH = f'{WORKDIR}/artwork_print_ready.jpg'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_print_requirements_doc():
    """Create a realistic print shop requirements document."""
    doc = Document()

    # Title
    title = doc.add_heading('Print Shop — Image Preparation Requirements', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph
    intro = doc.add_paragraph(
        'All artwork submitted for professional printing must meet the following technical '
        'specifications. Please ensure your file satisfies every requirement before submission. '
        'Files that do not comply will be returned for correction.'
    )

    doc.add_paragraph('')  # spacing

    # Section 1: Color Mode
    doc.add_heading('1. Color Mode', level=1)
    p = doc.add_paragraph()
    run = p.add_run('Required: ')
    run.bold = True
    p.add_run('RGB color mode')

    doc.add_paragraph(
        'The artwork must be in RGB color mode. Files in grayscale (L), RGBA, or palette '
        'mode (P) are not accepted. GIMP: Image > Mode > RGB.'
    )

    # Section 2: Resolution
    doc.add_heading('2. Print Resolution', level=1)
    p = doc.add_paragraph()
    run = p.add_run('Required: ')
    run.bold = True
    p.add_run('300 DPI (dots per inch)')

    doc.add_paragraph(
        'The image must be saved with a resolution of exactly 300 DPI to ensure sharp print '
        'quality. Low-resolution files (e.g., 72 DPI screen resolution) will appear blurry '
        'when printed. In GIMP, set via Image > Scale Image (resolution fields) or '
        'File > Export As with DPI metadata.'
    )

    # Section 3: Bleed Area
    doc.add_heading('3. Bleed Area', level=1)
    p = doc.add_paragraph()
    run = p.add_run('Required: ')
    run.bold = True
    p.add_run('3 mm bleed on all four sides')

    doc.add_paragraph(
        'A 3 mm bleed area is required on all four sides of the artwork. At 300 DPI, '
        '3 mm equals approximately 35 pixels. Add the bleed by extending the canvas in GIMP: '
        'Image > Canvas Size, then add 70 pixels to both Width and Height (35 px per side), '
        'center the existing layer, and fill the new border area with the background color '
        'or extend the edges. Flatten the image before exporting.'
    )

    # Section 4: File Format
    doc.add_heading('4. Output File', level=1)
    p = doc.add_paragraph()
    run = p.add_run('Save the print-ready file as: ')
    run.bold = True
    p.add_run('artwork_print_ready.jpg')

    doc.add_paragraph(
        'Export the prepared artwork to the Desktop as "artwork_print_ready.jpg" using '
        'JPEG format with quality 90 or higher. Do not rename the file.'
    )

    # Summary table
    doc.add_heading('Quick Reference Summary', level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = [('Requirement', 'Specification')]
    rows_data = [
        ('Color Mode', 'RGB'),
        ('Resolution', '300 DPI'),
        ('Bleed Area', '3 mm on all sides (~35 px at 300 DPI)'),
        ('Output Filename', 'artwork_print_ready.jpg'),
    ]

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Requirement'
    hdr_cells[1].text = 'Specification'
    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Data rows
    for i, (req, spec) in enumerate(rows_data, 1):
        table.rows[i].cells[0].text = req
        table.rows[i].cells[1].text = spec

    doc.save(REQUIREMENTS_DOC)
    print(f'Requirements document created: {REQUIREMENTS_DOC}')


def create_artwork():
    """Create a realistic artwork.jpg at 72 DPI (pre-print, screen resolution)."""
    # Create a 600x450 pixel image (approx A5 proportions at 72 dpi)
    width, height = 600, 450
    img = Image.new('RGB', (width, height), color=(245, 245, 245))
    draw = ImageDraw.Draw(img)

    # Background gradient-like fill using rectangles
    for i in range(height):
        r = int(30 + (i / height) * 80)
        g = int(60 + (i / height) * 100)
        b = int(120 + (i / height) * 80)
        draw.line([(0, i), (width, i)], fill=(r, g, b))

    # Decorative elements: sun/circle
    draw.ellipse([60, 40, 180, 160], fill=(255, 210, 50), outline=(220, 170, 20), width=3)

    # Mountains / triangles
    draw.polygon([(0, height), (150, 200), (300, height)], fill=(80, 100, 60))
    draw.polygon([(150, height), (320, 160), (500, height)], fill=(60, 85, 50))
    draw.polygon([(300, height), (460, 180), (width, height)], fill=(70, 95, 55))

    # Snow caps
    draw.polygon([(140, 210), (150, 200), (160, 210)], fill=(240, 240, 255))
    draw.polygon([(310, 170), (320, 160), (330, 170)], fill=(240, 240, 255))
    draw.polygon([(450, 190), (460, 180), (470, 190)], fill=(240, 240, 255))

    # River / path
    draw.polygon(
        [(220, height), (240, 360), (260, 330), (280, 360), (300, height)],
        fill=(100, 160, 210)
    )

    # Trees
    for tx in [80, 130, 430, 480, 530]:
        draw.polygon([(tx, 350), (tx - 20, 395), (tx + 20, 395)], fill=(40, 80, 40))
        draw.rectangle([tx - 5, 395, tx + 5, 420], fill=(90, 60, 30))

    # Border
    draw.rectangle([0, 0, width - 1, height - 1], outline=(30, 30, 30), width=3)

    # Text label
    try:
        font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf', 22)
        small_font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 14)
    except OSError:
        font = ImageFont.load_default()
        small_font = font

    draw.text((width // 2 - 80, 15), 'Mountain Vista', fill=(255, 255, 255), font=font)
    draw.text((width // 2 - 70, 42), 'Original Artwork — Screen Resolution', fill=(200, 200, 200), font=small_font)

    # Save at 72 DPI (screen resolution — needs print prep)
    img.save(ARTWORK_PATH, 'JPEG', quality=92, dpi=(72, 72))
    print(f'Artwork created: {ARTWORK_PATH} (600x450 px, 72 DPI)')


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    # Remove any pre-existing print-ready file (initial state must not have it)
    if os.path.exists(PRINT_READY_PATH):
        os.remove(PRINT_READY_PATH)
        print(f'Removed pre-existing file: {PRINT_READY_PATH}')

    # Create files
    create_print_requirements_doc()
    create_artwork()

    print('All initial files created.')

    # GUI-ready: Open the requirements document in LibreOffice Writer and artwork in GIMP
    # Open requirements doc first
    launch_gui(f'libreoffice --writer "{REQUIREMENTS_DOC}"', delay_sec=3.0)
    # Open artwork in GIMP
    launch_gui(f'gimp "{ARTWORK_PATH}"', delay_sec=2.0)

    print('GUI_READY: Launched LibreOffice Writer (print_requirements.docx) and GIMP (artwork.jpg) with DISPLAY=:0')


create_initial()
