"""
Reward Script: Insert a caption below the screenshot image on page 4
Task ID: writer_tech_030
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Caption paragraph with correct text exists
  Component 2 (0.3): Caption is positioned immediately after the image paragraph
  Component 3 (0.3): Caption uses proper Writer caption style (Caption style, italic, centered)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_030'


def persist_app_state(domain):
    """Best-effort save of any open LibreOffice document."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        import time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not installed")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------------------------------------------------------------
    # Component 1: Caption paragraph with correct text exists (0.4 points)
    # Task requires: 'Figure 1: Dashboard Overview'
    # This FAILS on initial (no such text), PASSES on golden.
    # ---------------------------------------------------------------
    try:
        caption_text_target = "Figure 1: Dashboard Overview"
        caption_para_idx = next(
            (i for i, para in enumerate(doc.paragraphs)
             if caption_text_target.lower() in para.text.strip().lower()),
            None
        )

        if caption_para_idx is not None:
            print(f"PASS: Component 1 — Caption text '{caption_text_target}' found at paragraph {caption_para_idx} (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Caption text '{caption_text_target}' not found in any paragraph")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------------------------------------------------------------
    # Component 2: Caption is positioned after the image paragraph (0.3 points)
    # The image (w:drawing element) should be in a paragraph before the caption.
    # This FAILS on initial (no caption at all), PASSES on golden.
    # ---------------------------------------------------------------
    try:
        if caption_para_idx is not None:
            # Look for an image paragraph before the caption
            image_para_idx = None
            # Search backwards from caption for nearest paragraph containing a drawing
            for j in range(caption_para_idx - 1, max(caption_para_idx - 5, -1), -1):
                if j < 0:
                    break
                para_xml = doc.paragraphs[j]._element.xml
                if 'w:drawing' in para_xml or 'graphicData' in para_xml:
                    image_para_idx = j
                    break

            if image_para_idx is not None:
                # The caption should be within 1-2 paragraphs after the image
                gap = caption_para_idx - image_para_idx
                if gap <= 2:
                    print(f"PASS: Component 2 — Caption at P{caption_para_idx} is {gap} paragraph(s) after image at P{image_para_idx} (0.3 pts)")
                    total_score += 0.3
                else:
                    print(f"FAIL: Component 2 — Caption at P{caption_para_idx} is {gap} paragraphs after image at P{image_para_idx} (too far)")
            else:
                print(f"FAIL: Component 2 — No image paragraph found before caption at P{caption_para_idx}")
        else:
            print(f"FAIL: Component 2 — No caption found, cannot check position")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---------------------------------------------------------------
    # Component 3: Caption uses proper Writer caption formatting (0.3 points)
    # Standard Writer caption: 'Caption' style, italic text, centered alignment.
    # This FAILS on initial (no caption), PASSES on golden.
    # ---------------------------------------------------------------
    try:
        if caption_para_idx is not None:
            cap_para = doc.paragraphs[caption_para_idx]
            style_name = cap_para.style.name if cap_para.style else "None"
            alignment = cap_para.paragraph_format.alignment

            # Sub-checks for formatting
            style_ok = style_name.lower() == 'caption'
            # Check italic on runs
            italic_ok = False
            if cap_para.runs:
                italic_ok = all(
                    r.italic or r.font.italic
                    for r in cap_para.runs
                    if r.text.strip()
                )

            sub_score = 0.0
            details = []

            # Style check (0.15)
            if style_ok:
                sub_score += 0.15
                details.append(f"style='Caption' OK")
            else:
                details.append(f"style='{style_name}' (expected 'Caption')")

            # Italic check (0.075)
            if italic_ok:
                sub_score += 0.075
                details.append("italic OK")
            else:
                details.append("italic MISSING")

            # Centered alignment check (0.075)
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            centered = (alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
            if centered:
                sub_score += 0.075
                details.append("centered OK")
            else:
                details.append(f"alignment={alignment} (expected CENTER)")

            if sub_score > 0:
                print(f"PASS: Component 3 — Caption formatting: {', '.join(details)} ({sub_score} pts)")
                total_score += sub_score
            else:
                print(f"FAIL: Component 3 — Caption formatting incorrect: {', '.join(details)}")
        else:
            print(f"FAIL: Component 3 — No caption found, cannot check formatting")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
