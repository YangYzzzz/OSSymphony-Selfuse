"""
Initial Setup: Insert a caption below the screenshot image on page 4
Task ID: writer_tech_030
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# Generate a simple dashboard screenshot image using PIL
from PIL import Image, ImageDraw, ImageFont

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_030'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
IMG_PATH = f'{WORKDIR}/dashboard_screenshot.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_dashboard_image():
    """Create a simple dashboard screenshot image."""
    img = Image.new('RGB', (800, 500), color=(240, 242, 245))
    draw = ImageDraw.Draw(img)

    # Title bar
    draw.rectangle([0, 0, 800, 50], fill=(33, 37, 41))
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 18)
        font_sm = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 12)
        font_lg = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 24)
    except Exception:
        font = ImageFont.load_default()
        font_sm = font
        font_lg = font

    draw.text((20, 14), "Analytics Dashboard", fill=(255, 255, 255), font=font)

    # KPI cards
    cards = [
        ("Total Users", "12,458", (52, 152, 219)),
        ("Revenue", "$84,230", (46, 204, 113)),
        ("Orders", "3,127", (155, 89, 182)),
        ("Growth", "+15.3%", (230, 126, 34)),
    ]
    for i, (label, value, color) in enumerate(cards):
        x = 20 + i * 195
        draw.rectangle([x, 70, x + 180, 150], fill=(255, 255, 255), outline=(200, 200, 200))
        draw.text((x + 15, 82), label, fill=(100, 100, 100), font=font_sm)
        draw.text((x + 15, 105), value, fill=color, font=font_lg)

    # Chart area placeholder
    draw.rectangle([20, 170, 500, 400], fill=(255, 255, 255), outline=(200, 200, 200))
    draw.text((200, 175), "Monthly Trends", fill=(50, 50, 50), font=font)
    # Simple bar chart
    bar_values = [120, 180, 150, 220, 200, 280, 250, 310, 270, 350, 320, 380]
    max_val = max(bar_values)
    for j, val in enumerate(bar_values):
        bx = 40 + j * 37
        bh = int((val / max_val) * 180)
        draw.rectangle([bx, 390 - bh, bx + 28, 390], fill=(52, 152, 219))

    # Side panel
    draw.rectangle([520, 170, 780, 400], fill=(255, 255, 255), outline=(200, 200, 200))
    draw.text((540, 180), "Top Products", fill=(50, 50, 50), font=font)
    products = [
        ("Enterprise Suite", "$24,500"),
        ("Cloud Platform", "$18,320"),
        ("API Gateway", "$12,890"),
        ("Data Analytics", "$9,740"),
        ("Security Shield", "$7,650"),
    ]
    for k, (name, rev) in enumerate(products):
        y = 210 + k * 35
        draw.text((540, y), name, fill=(70, 70, 70), font=font_sm)
        draw.text((700, y), rev, fill=(46, 204, 113), font=font_sm)

    # Footer
    draw.rectangle([0, 470, 800, 500], fill=(248, 249, 250))
    draw.text((20, 478), "Last updated: March 28, 2026  |  Data source: Internal Analytics Platform", fill=(130, 130, 130), font=font_sm)

    img.save(IMG_PATH)
    print(f'Dashboard image created: {IMG_PATH}')


def create_initial():
    create_dashboard_image()

    doc = Document()

    # --- Page 1: Title and Introduction ---
    title = doc.add_heading('Quarterly Performance Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('TechVision Analytics — Q1 2026')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacing

    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This report provides a comprehensive overview of TechVision Analytics\' '
        'performance during the first quarter of 2026. Key highlights include '
        'sustained revenue growth of 15.3% year-over-year, expansion of the customer '
        'base to over 12,000 active accounts, and successful launch of three new '
        'product features that have driven significant engagement improvements.'
    )
    doc.add_paragraph(
        'The Engineering team completed the migration to the new microservices '
        'architecture ahead of schedule, resulting in a 40% improvement in API '
        'response times. Customer satisfaction scores reached an all-time high of '
        '4.7 out of 5.0, reflecting the team\'s commitment to product quality.'
    )

    # --- Page 2: Financial Overview ---
    doc.add_page_break()

    doc.add_heading('2. Financial Overview', level=1)
    doc.add_paragraph(
        'Total revenue for Q1 2026 reached $84,230, representing a 15.3% increase '
        'compared to Q1 2025. The growth was primarily driven by enterprise subscription '
        'upgrades and new customer acquisitions in the APAC region.'
    )

    doc.add_heading('2.1 Revenue Breakdown by Product', level=2)

    # Revenue table
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    headers = ['Product', 'Q1 2025', 'Q1 2026', 'Growth']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    revenue_data = [
        ['Enterprise Suite', '$20,100', '$24,500', '+21.9%'],
        ['Cloud Platform', '$15,200', '$18,320', '+20.5%'],
        ['API Gateway', '$11,500', '$12,890', '+12.1%'],
        ['Data Analytics', '$8,900', '$9,740', '+9.4%'],
        ['Security Shield', '$6,300', '$7,650', '+21.4%'],
        ['Other Services', '$10,080', '$11,130', '+10.4%'],
    ]
    for r, row_data in enumerate(revenue_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_paragraph(
        'The Enterprise Suite continues to be our flagship product, contributing 29.1% '
        'of total revenue. The Cloud Platform showed the strongest absolute growth, '
        'adding $3,120 in incremental revenue compared to the prior year.'
    )

    # --- Page 3: Operational Metrics ---
    doc.add_page_break()

    doc.add_heading('3. Operational Metrics', level=1)

    doc.add_heading('3.1 Customer Acquisition', level=2)
    doc.add_paragraph(
        'New customer sign-ups totaled 1,847 during Q1, a 23% increase from the previous '
        'quarter. The sales team successfully closed 42 enterprise deals, with an average '
        'contract value of $18,500. Customer churn rate decreased to 2.1%, down from 3.4% '
        'in Q4 2025.'
    )

    doc.add_heading('3.2 Engineering Performance', level=2)
    doc.add_paragraph(
        'The engineering team deployed 127 releases during the quarter, maintaining a '
        '99.97% uptime SLA. Key infrastructure improvements included:'
    )
    items = [
        'Migration to Kubernetes-based microservices architecture',
        'Implementation of real-time data streaming pipeline',
        'Launch of automated security scanning for CI/CD pipeline',
        'Deployment of edge caching nodes in 12 new regions',
        'Reduction of average API latency from 145ms to 87ms',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.3 Support Metrics', level=2)

    support_table = doc.add_table(rows=5, cols=3)
    support_table.style = 'Table Grid'
    s_headers = ['Metric', 'Q4 2025', 'Q1 2026']
    for i, h in enumerate(s_headers):
        run = support_table.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True

    support_data = [
        ['Avg. Response Time', '4.2 hours', '2.8 hours'],
        ['Resolution Rate', '94.1%', '96.7%'],
        ['CSAT Score', '4.5/5.0', '4.7/5.0'],
        ['Tickets Handled', '3,412', '3,890'],
    ]
    for r, row_data in enumerate(support_data, 1):
        for c, val in enumerate(row_data):
            support_table.cell(r, c).text = val

    # --- Page 4: Dashboard & Visual Analytics ---
    doc.add_page_break()

    doc.add_heading('4. Dashboard & Visual Analytics', level=1)
    doc.add_paragraph(
        'The following screenshot captures the current state of our internal analytics '
        'dashboard, providing a real-time view of key performance indicators, monthly '
        'revenue trends, and top-performing product lines.'
    )

    # Insert the dashboard screenshot image - NO CAPTION
    doc.add_picture(IMG_PATH, width=Inches(5.5))
    # Center the image
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()
    doc.add_paragraph(
        'The dashboard demonstrates consistent upward trends across all primary metrics. '
        'Monthly revenue has shown steady growth since January, with March recording '
        'the highest single-month revenue in company history at $31,200.'
    )

    # --- Page 5: Outlook ---
    doc.add_page_break()

    doc.add_heading('5. Q2 2026 Outlook', level=1)
    doc.add_paragraph(
        'Looking ahead to Q2 2026, we anticipate continued growth momentum driven by '
        'several strategic initiatives:'
    )
    outlook_items = [
        'Launch of the AI-powered recommendation engine for enterprise clients',
        'Expansion into Latin American markets with localized product offerings',
        'Release of the new mobile SDK for iOS and Android platforms',
        'Partnership with three major cloud providers for integrated deployments',
    ]
    for item in outlook_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'Revenue projections for Q2 range between $92,000 and $97,000, representing '
        'a potential quarter-over-quarter growth of 9-15%. The team remains focused on '
        'execution excellence and customer value delivery.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
