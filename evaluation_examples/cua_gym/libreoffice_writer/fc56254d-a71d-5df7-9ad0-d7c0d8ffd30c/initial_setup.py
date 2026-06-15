"""
Initial Setup: Resource guide with 12 hyperlinks to be removed
Task ID: osworld_writer_easy_032
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_easy_032'
OUTPUT = f'{WORKDIR}/resource_guide.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink run to a paragraph in python-docx.
    Returns the hyperlink element.
    """
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')

    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Community Technology Resource Guide', level=0)

    intro = doc.add_paragraph(
        'This guide provides a comprehensive overview of digital tools, online platforms, '
        'and technology resources available to our community members. Whether you are '
        'looking to improve your digital skills, find remote work opportunities, or access '
        'government services online, the resources listed here will help you get started.'
    )

    # --- Section 1: Digital Literacy ---
    doc.add_heading('1. Digital Literacy and Training', level=1)

    p1 = doc.add_paragraph(
        'Building foundational digital skills is essential in today\'s connected world. '
        'The GCFGlobal platform offers free courses on everything from basic computer '
        'use to advanced spreadsheet techniques. Visit their learning portal at '
    )
    add_hyperlink(p1, 'https://edu.gcfglobal.org/en/', 'https://edu.gcfglobal.org/en/')
    p1.add_run(' to browse available courses.')

    p2 = doc.add_paragraph(
        'For more structured learning paths, Coursera provides both free and paid '
        'certifications in technology fields. Community members seeking financial aid '
        'can apply through the assistance program — '
    )
    add_hyperlink(p2, 'https://www.coursera.org/financial-aid', 'click here')
    p2.add_run(' to access the application form.')

    p3 = doc.add_paragraph(
        'The Khan Academy also maintains a robust technology curriculum suitable for '
        'learners of all ages. Their computing section is available at '
    )
    add_hyperlink(p3, 'https://www.khanacademy.org/computing', 'https://www.khanacademy.org/computing')
    p3.add_run('.')

    # --- Section 2: Remote Work Resources ---
    doc.add_heading('2. Remote Work and Employment', level=1)

    p4 = doc.add_paragraph(
        'Remote work opportunities have expanded significantly. LinkedIn remains the '
        'leading professional network for job seekers; create or update your profile at '
    )
    add_hyperlink(p4, 'https://www.linkedin.com/jobs/', 'https://www.linkedin.com/jobs/')
    p4.add_run(' to connect with employers.')

    p5 = doc.add_paragraph(
        'Upwork and similar freelance platforms connect independent contractors with '
        'clients worldwide. To register as a freelancer and start earning, '
    )
    add_hyperlink(p5, 'https://www.upwork.com/freelance-jobs/', 'learn more here')
    p5.add_run('.')

    p6 = doc.add_paragraph(
        'The US Department of Labor\'s CareerOneStop website offers job search tools, '
        'resume builders, and training resources. Access the full toolkit at '
    )
    add_hyperlink(p6, 'https://www.careeronestop.org/', 'https://www.careeronestop.org/')
    p6.add_run('.')

    # --- Section 3: Government Services ---
    doc.add_heading('3. Government and Social Services Online', level=1)

    p7 = doc.add_paragraph(
        'Benefits.gov is the official U.S. government website for benefit programs. '
        'You can check your eligibility for over 1,000 federal and state benefit programs at '
    )
    add_hyperlink(p7, 'https://www.benefits.gov/', 'https://www.benefits.gov/')
    p7.add_run('.')

    p8 = doc.add_paragraph(
        'For healthcare coverage under the Affordable Care Act, the HealthCare.gov '
        'marketplace allows individuals to compare plans and enroll. '
    )
    add_hyperlink(p8, 'https://www.healthcare.gov/get-coverage/', 'Get coverage now')
    p8.add_run(' by visiting the enrollment portal.')

    p9 = doc.add_paragraph(
        'The IRS Free File program allows eligible taxpayers to prepare and file '
        'federal income tax returns for free. Eligible filers can access the program at '
    )
    add_hyperlink(p9, 'https://www.irs.gov/filing/free-file-do-your-federal-taxes-for-free', 'https://www.irs.gov/filing/free-file-do-your-federal-taxes-for-free')
    p9.add_run('.')

    # --- Section 4: Health and Wellness ---
    doc.add_heading('4. Health and Wellness Resources', level=1)

    p10 = doc.add_paragraph(
        'MedlinePlus, operated by the National Library of Medicine, provides reliable '
        'health information in both English and Spanish. Browse topics and articles at '
    )
    add_hyperlink(p10, 'https://medlineplus.gov/', 'https://medlineplus.gov/')
    p10.add_run('.')

    p11 = doc.add_paragraph(
        'Crisis support is available 24/7 through the 988 Suicide and Crisis Lifeline. '
        'For additional mental health resources and to find local providers, '
    )
    add_hyperlink(p11, 'https://www.samhsa.gov/find-help/national-helpline', 'visit SAMHSA\'s helpline page')
    p11.add_run('.')

    # --- Section 5: Community ---
    doc.add_heading('5. Community and Civic Engagement', level=1)

    p12 = doc.add_paragraph(
        'VolunteerMatch helps connect passionate people with nonprofits that need help. '
        'Find volunteer opportunities in your area at '
    )
    add_hyperlink(p12, 'https://www.volunteermatch.org/', 'https://www.volunteermatch.org/')
    p12.add_run('. Your skills and time can make a real difference in your community.')

    # --- Footer note ---
    doc.add_paragraph('')
    note = doc.add_paragraph(
        'Note: This resource guide is updated quarterly. If you find a broken link or '
        'have a resource to suggest, please contact the community outreach coordinator '
        'at outreach@communitycenter.org.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
