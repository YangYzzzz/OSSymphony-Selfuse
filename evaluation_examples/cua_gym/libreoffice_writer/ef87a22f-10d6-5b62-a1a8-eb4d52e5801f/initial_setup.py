"""
Initial Setup: Legal contract document with exhibits - all in default portrait style
Task ID: writer_legal_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_043'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    return h


def add_body_para(doc, text, bold=False, alignment=None, space_after=Pt(6)):
    """Add a body paragraph with optional styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    if alignment:
        p.paragraph_format.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p


def create_initial():
    doc = Document()

    # Set default page style - portrait, 1 inch margins
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # === TITLE PAGE ===
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('PROFESSIONAL SERVICES AGREEMENT')
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Between')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    party1 = doc.add_paragraph()
    party1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = party1.add_run('MERIDIAN TECHNOLOGY SOLUTIONS, INC.')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    and_para = doc.add_paragraph()
    and_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = and_para.add_run('and')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    party2 = doc.add_paragraph()
    party2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = party2.add_run('CASCADIA HEALTHCARE GROUP, LLC')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Effective Date: January 15, 2026')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    contract_num = doc.add_paragraph()
    contract_num.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = contract_num.add_run('Contract No. MTS-CHG-2026-0847')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # === PAGE BREAK - RECITALS ===
    doc.add_page_break()

    add_heading_styled(doc, 'RECITALS', level=1)

    add_body_para(doc, 'WHEREAS, Meridian Technology Solutions, Inc. ("Provider") is a Delaware corporation '
                  'engaged in the business of providing information technology consulting, software development, '
                  'and managed services to healthcare organizations throughout the United States;')

    add_body_para(doc, 'WHEREAS, Cascadia Healthcare Group, LLC ("Client") is a Washington limited liability '
                  'company operating a network of fourteen (14) regional hospitals, twenty-three (23) outpatient '
                  'clinics, and seven (7) specialized rehabilitation centers across the Pacific Northwest;')

    add_body_para(doc, 'WHEREAS, Client desires to engage Provider to perform certain professional services '
                  'relating to the modernization and integration of Client\'s electronic health record ("EHR") '
                  'systems, patient data management platforms, and associated clinical workflow applications;')

    add_body_para(doc, 'WHEREAS, Provider has represented that it possesses the requisite expertise, personnel, '
                  'and resources to perform the services described herein in accordance with applicable healthcare '
                  'industry standards, including but not limited to HIPAA, HITECH, and state privacy regulations;')

    add_body_para(doc, 'NOW, THEREFORE, in consideration of the mutual covenants and agreements hereinafter set '
                  'forth and for other good and valuable consideration, the receipt and sufficiency of which are '
                  'hereby acknowledged, the parties agree as follows:')

    # === ARTICLE I - DEFINITIONS ===
    doc.add_page_break()
    add_heading_styled(doc, 'ARTICLE I — DEFINITIONS', level=1)

    definitions = [
        ('"Acceptance Criteria"', 'means the specific technical and functional requirements that Deliverables must satisfy prior to Client\'s acceptance, as detailed in the applicable Statement of Work.'),
        ('"Authorized Users"', 'means employees, contractors, and agents of Client who are authorized to access and use the Deliverables in connection with Client\'s business operations.'),
        ('"Change Order"', 'means a written amendment to a Statement of Work, executed by both parties, that modifies the scope, timeline, or fees for Services under such Statement of Work.'),
        ('"Confidential Information"', 'means any non-public information disclosed by either party, including but not limited to trade secrets, business plans, financial data, patient information, technical specifications, and proprietary algorithms.'),
        ('"Deliverables"', 'means the tangible and intangible work product, including software, documentation, reports, and other materials, produced by Provider in the performance of the Services.'),
        ('"Effective Date"', 'means January 15, 2026, the date on which this Agreement becomes legally binding upon both parties.'),
        ('"Force Majeure Event"', 'means any event beyond the reasonable control of the affected party, including natural disasters, pandemics, government actions, cyber attacks, or utility failures.'),
        ('"Intellectual Property Rights"', 'means all patents, copyrights, trademarks, trade secrets, and other proprietary rights recognized under applicable law.'),
    ]

    for term, definition in definitions:
        p = doc.add_paragraph()
        run_term = p.add_run(f'1.{definitions.index((term, definition)) + 1}  {term} ')
        run_term.bold = True
        run_term.font.name = 'Times New Roman'
        run_term.font.size = Pt(12)
        run_def = p.add_run(definition)
        run_def.font.name = 'Times New Roman'
        run_def.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    # === ARTICLE II - SCOPE OF SERVICES ===
    doc.add_page_break()
    add_heading_styled(doc, 'ARTICLE II — SCOPE OF SERVICES', level=1)

    add_body_para(doc, '2.1  General Scope. Provider shall perform the professional services described in each '
                  'Statement of Work executed pursuant to this Agreement (collectively, the "Services"). The initial '
                  'Statement of Work, attached hereto as Exhibit A, describes the Phase I implementation services.')

    add_body_para(doc, '2.2  Standards of Performance. Provider shall perform all Services in a professional and '
                  'workmanlike manner, consistent with generally accepted industry standards for healthcare IT '
                  'consulting services. Provider shall assign qualified personnel with appropriate certifications '
                  'and experience to perform the Services.')

    add_body_para(doc, '2.3  Compliance Requirements. Provider acknowledges that Client operates in a heavily '
                  'regulated industry and agrees to comply with all applicable federal, state, and local laws '
                  'and regulations, including without limitation the Health Insurance Portability and Accountability '
                  'Act of 1996 ("HIPAA"), the Health Information Technology for Economic and Clinical Health Act '
                  '("HITECH"), and the Washington State Patient Privacy Act.')

    add_body_para(doc, '2.4  Project Governance. The parties shall establish a joint steering committee consisting '
                  'of two (2) representatives from each party. The steering committee shall meet bi-weekly to review '
                  'project progress, address issues, and approve Change Orders. Decisions of the steering committee '
                  'shall be documented in meeting minutes distributed within three (3) business days.')

    add_body_para(doc, '2.5  Subcontracting. Provider shall not subcontract any portion of the Services without '
                  'prior written consent of Client. Any approved subcontractor shall be subject to confidentiality '
                  'obligations no less restrictive than those set forth in this Agreement.')

    # === ARTICLE III - COMPENSATION ===
    doc.add_page_break()
    add_heading_styled(doc, 'ARTICLE III — COMPENSATION AND PAYMENT', level=1)

    add_body_para(doc, '3.1  Fees. Client shall pay Provider the fees set forth in each Statement of Work. Unless '
                  'otherwise specified, fees shall be calculated on a time-and-materials basis at the hourly rates '
                  'specified in Exhibit B (Rate Schedule).')

    add_body_para(doc, '3.2  Invoicing. Provider shall submit detailed monthly invoices to Client no later than the '
                  'tenth (10th) business day following the end of each calendar month. Each invoice shall include: '
                  '(a) a description of Services performed; (b) the name and title of each Provider personnel who '
                  'performed Services; (c) the number of hours worked by each such person; (d) the applicable hourly '
                  'rate; and (e) any pre-approved expenses.')

    add_body_para(doc, '3.3  Payment Terms. Client shall pay all undisputed invoice amounts within thirty (30) days '
                  'of receipt. Late payments shall bear interest at the rate of one and one-half percent (1.5%) per '
                  'month, or the maximum rate permitted by law, whichever is less.')

    add_body_para(doc, '3.4  Expense Reimbursement. Client shall reimburse Provider for reasonable, pre-approved '
                  'travel and out-of-pocket expenses incurred in connection with the Services. All expenses exceeding '
                  'Five Hundred Dollars ($500.00) require prior written approval from Client\'s project manager.')

    add_body_para(doc, '3.5  Taxes. Fees are exclusive of all applicable taxes. Client shall be responsible for all '
                  'sales, use, and similar taxes arising from this Agreement, excluding taxes based on Provider\'s '
                  'net income.')

    # === ARTICLE IV - INTELLECTUAL PROPERTY ===
    doc.add_page_break()
    add_heading_styled(doc, 'ARTICLE IV — INTELLECTUAL PROPERTY RIGHTS', level=1)

    add_body_para(doc, '4.1  Work Product Ownership. All Deliverables created by Provider specifically for Client '
                  'under this Agreement shall be considered "work made for hire" to the maximum extent permitted by '
                  'law. To the extent any Deliverable does not qualify as work made for hire, Provider hereby assigns '
                  'to Client all right, title, and interest in such Deliverable.')

    add_body_para(doc, '4.2  Provider Pre-Existing IP. Provider retains all rights in its pre-existing intellectual '
                  'property, including tools, methodologies, frameworks, and code libraries that existed prior to this '
                  'Agreement ("Provider IP"). To the extent Provider IP is incorporated into any Deliverable, Provider '
                  'grants Client a non-exclusive, perpetual, royalty-free license to use such Provider IP solely in '
                  'connection with Client\'s use of the Deliverable.')

    add_body_para(doc, '4.3  Open Source Components. Provider shall identify all open source software components '
                  'incorporated into Deliverables prior to delivery. Provider shall not incorporate any open source '
                  'component with a copyleft license (e.g., GPL, AGPL) without Client\'s prior written consent.')

    add_body_para(doc, '4.4  Feedback. Client grants Provider a non-exclusive, perpetual license to use any feedback, '
                  'suggestions, or ideas provided by Client regarding the Services or Deliverables to improve '
                  'Provider\'s general products and services, provided such use does not disclose Client\'s '
                  'Confidential Information.')

    # === ARTICLE V - CONFIDENTIALITY ===
    doc.add_page_break()
    add_heading_styled(doc, 'ARTICLE V — CONFIDENTIALITY', level=1)

    add_body_para(doc, '5.1  Obligations. Each party agrees to: (a) hold the other party\'s Confidential Information '
                  'in strict confidence; (b) not disclose such information to any third party except as expressly '
                  'permitted herein; (c) use such information solely for the purpose of performing obligations under '
                  'this Agreement; and (d) protect such information using at least the same degree of care it uses '
                  'to protect its own confidential information, but in no event less than reasonable care.')

    add_body_para(doc, '5.2  Exceptions. Confidential Information does not include information that: (a) is or becomes '
                  'publicly available through no breach of this Agreement; (b) was rightfully in the receiving party\'s '
                  'possession prior to disclosure; (c) is independently developed without use of the disclosing party\'s '
                  'Confidential Information; or (d) is rightfully received from a third party without restriction.')

    add_body_para(doc, '5.3  Required Disclosures. A receiving party may disclose Confidential Information to the extent '
                  'required by applicable law, regulation, or court order, provided that the receiving party gives the '
                  'disclosing party prompt written notice and reasonable cooperation to seek protective measures.')

    add_body_para(doc, '5.4  Return of Materials. Upon termination of this Agreement or upon request, each party shall '
                  'promptly return or destroy all tangible materials containing the other party\'s Confidential '
                  'Information and certify such return or destruction in writing.')

    # === ARTICLE VI - TERM AND TERMINATION ===
    add_heading_styled(doc, 'ARTICLE VI — TERM AND TERMINATION', level=1)

    add_body_para(doc, '6.1  Term. This Agreement shall commence on the Effective Date and continue for a period of '
                  'three (3) years, unless earlier terminated in accordance with this Article VI. The Agreement shall '
                  'automatically renew for successive one (1) year periods unless either party provides written notice '
                  'of non-renewal at least ninety (90) days prior to the end of the then-current term.')

    add_body_para(doc, '6.2  Termination for Convenience. Either party may terminate this Agreement for any reason by '
                  'providing sixty (60) days\' prior written notice to the other party.')

    add_body_para(doc, '6.3  Termination for Cause. Either party may terminate this Agreement immediately upon written '
                  'notice if the other party: (a) materially breaches this Agreement and fails to cure such breach '
                  'within thirty (30) days of written notice; (b) becomes insolvent or files for bankruptcy; or '
                  '(c) is convicted of fraud or other criminal conduct related to this Agreement.')

    add_body_para(doc, '6.4  Effect of Termination. Upon termination: (a) Client shall pay Provider for all Services '
                  'performed and expenses incurred through the effective date of termination; (b) Provider shall deliver '
                  'all completed and in-progress Deliverables; and (c) the provisions of Articles IV, V, VII, and IX '
                  'shall survive termination.')

    # === ARTICLE VII - WARRANTIES ===
    doc.add_page_break()
    add_heading_styled(doc, 'ARTICLE VII — REPRESENTATIONS AND WARRANTIES', level=1)

    add_body_para(doc, '7.1  Provider Warranties. Provider represents and warrants that: (a) it has the legal right '
                  'and authority to enter into this Agreement; (b) the Services will be performed in a professional '
                  'manner; (c) the Deliverables will substantially conform to the specifications in the applicable '
                  'Statement of Work; and (d) the Deliverables will not infringe any third-party intellectual property rights.')

    add_body_para(doc, '7.2  Client Warranties. Client represents and warrants that: (a) it has the legal right and '
                  'authority to enter into this Agreement; (b) it will provide reasonable cooperation and access to '
                  'systems and personnel as necessary for Provider to perform the Services; and (c) all information '
                  'provided to Provider is accurate and complete to the best of Client\'s knowledge.')

    add_body_para(doc, '7.3  Disclaimer. EXCEPT AS EXPRESSLY SET FORTH IN THIS AGREEMENT, PROVIDER MAKES NO OTHER '
                  'WARRANTIES, EXPRESS OR IMPLIED, INCLUDING WITHOUT LIMITATION ANY IMPLIED WARRANTIES OF MERCHANTABILITY '
                  'OR FITNESS FOR A PARTICULAR PURPOSE.')

    # === ARTICLE VIII - LIMITATION OF LIABILITY ===
    add_heading_styled(doc, 'ARTICLE VIII — LIMITATION OF LIABILITY', level=1)

    add_body_para(doc, '8.1  Cap on Liability. EXCEPT FOR BREACHES OF ARTICLE V (CONFIDENTIALITY) AND CLAIMS ARISING '
                  'UNDER ARTICLE IX (INDEMNIFICATION), NEITHER PARTY\'S AGGREGATE LIABILITY UNDER THIS AGREEMENT SHALL '
                  'EXCEED THE TOTAL FEES PAID OR PAYABLE BY CLIENT UNDER THIS AGREEMENT DURING THE TWELVE (12) MONTH '
                  'PERIOD IMMEDIATELY PRECEDING THE EVENT GIVING RISE TO THE CLAIM.')

    add_body_para(doc, '8.2  Exclusion of Damages. IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, '
                  'SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING WITHOUT LIMITATION LOST PROFITS, LOST DATA, '
                  'OR BUSINESS INTERRUPTION, REGARDLESS OF THE FORM OF ACTION AND WHETHER OR NOT SUCH DAMAGES WERE '
                  'FORESEEABLE.')

    # === ARTICLE IX - INDEMNIFICATION ===
    add_heading_styled(doc, 'ARTICLE IX — INDEMNIFICATION', level=1)

    add_body_para(doc, '9.1  Provider Indemnification. Provider shall defend, indemnify, and hold harmless Client and its '
                  'officers, directors, employees, and agents from and against all claims, losses, damages, liabilities, '
                  'costs, and expenses (including reasonable attorneys\' fees) arising from: (a) Provider\'s breach of this '
                  'Agreement; (b) Provider\'s negligent or willful misconduct; or (c) any claim that the Deliverables '
                  'infringe a third party\'s intellectual property rights.')

    add_body_para(doc, '9.2  Client Indemnification. Client shall defend, indemnify, and hold harmless Provider from '
                  'claims arising from: (a) Client\'s breach of this Agreement; (b) Client\'s negligent or willful '
                  'misconduct; or (c) Client\'s unauthorized modification or misuse of the Deliverables.')

    # === ARTICLE X - GENERAL PROVISIONS ===
    doc.add_page_break()
    add_heading_styled(doc, 'ARTICLE X — GENERAL PROVISIONS', level=1)

    add_body_para(doc, '10.1  Governing Law. This Agreement shall be governed by and construed in accordance with the '
                  'laws of the State of Washington, without regard to its conflict of laws provisions.')

    add_body_para(doc, '10.2  Dispute Resolution. Any dispute arising under this Agreement shall first be submitted to '
                  'good faith mediation. If mediation is unsuccessful within thirty (30) days, either party may pursue '
                  'binding arbitration in Seattle, Washington, under the rules of the American Arbitration Association.')

    add_body_para(doc, '10.3  Notices. All notices under this Agreement shall be in writing and delivered by certified '
                  'mail, overnight courier, or electronic mail with confirmation of receipt to the addresses set forth '
                  'on the signature page.')

    add_body_para(doc, '10.4  Assignment. Neither party may assign this Agreement without the prior written consent of '
                  'the other party, except that either party may assign this Agreement in connection with a merger, '
                  'acquisition, or sale of substantially all of its assets.')

    add_body_para(doc, '10.5  Entire Agreement. This Agreement, together with all Statements of Work, Exhibits, and '
                  'Change Orders, constitutes the entire agreement between the parties and supersedes all prior '
                  'negotiations, representations, and agreements.')

    add_body_para(doc, '10.6  Amendments. This Agreement may be amended only by a written instrument executed by '
                  'authorized representatives of both parties.')

    add_body_para(doc, '10.7  Severability. If any provision of this Agreement is held invalid or unenforceable, the '
                  'remaining provisions shall continue in full force and effect.')

    add_body_para(doc, '10.8  Waiver. The failure of either party to enforce any provision of this Agreement shall not '
                  'constitute a waiver of such provision or any other provision.')

    # === SIGNATURE PAGE ===
    doc.add_page_break()
    add_heading_styled(doc, 'SIGNATURE PAGE', level=1)

    add_body_para(doc, 'IN WITNESS WHEREOF, the parties have executed this Professional Services Agreement as of the '
                  'Effective Date first written above.')

    doc.add_paragraph()

    add_body_para(doc, 'MERIDIAN TECHNOLOGY SOLUTIONS, INC.', bold=True)
    doc.add_paragraph()
    add_body_para(doc, 'By: ________________________________')
    add_body_para(doc, 'Name: Victoria R. Harrington')
    add_body_para(doc, 'Title: Chief Executive Officer')
    add_body_para(doc, 'Date: January 15, 2026')

    doc.add_paragraph()
    doc.add_paragraph()

    add_body_para(doc, 'CASCADIA HEALTHCARE GROUP, LLC', bold=True)
    doc.add_paragraph()
    add_body_para(doc, 'By: ________________________________')
    add_body_para(doc, 'Name: Dr. Jonathan M. Blackwell')
    add_body_para(doc, 'Title: Managing Partner')
    add_body_para(doc, 'Date: January 15, 2026')

    # === EXHIBIT A (starting around page 12) ===
    doc.add_page_break()

    add_heading_styled(doc, 'EXHIBIT A', level=1)
    add_heading_styled(doc, 'STATEMENT OF WORK — PHASE I: EHR SYSTEM MODERNIZATION', level=2)

    add_body_para(doc, 'This Statement of Work ("SOW") is entered into pursuant to the Professional Services Agreement '
                  'dated January 15, 2026, between Meridian Technology Solutions, Inc. ("Provider") and Cascadia '
                  'Healthcare Group, LLC ("Client").')

    add_body_para(doc, '1. PROJECT OVERVIEW', bold=True)
    add_body_para(doc, 'Provider shall perform a comprehensive modernization of Client\'s electronic health record '
                  'system, including migration from the legacy MedTrack 4.2 platform to the CloudHealth Enterprise '
                  'Suite version 8.0. The project encompasses data migration, system configuration, custom module '
                  'development, user training, and post-deployment support.')

    add_body_para(doc, '2. SCOPE OF WORK', bold=True)
    add_body_para(doc, '2.1  Discovery and Assessment Phase (Weeks 1-4): Provider shall conduct a thorough analysis of '
                  'Client\'s existing IT infrastructure, including interviews with key stakeholders, documentation '
                  'review, and technical system assessments across all fourteen hospital facilities.')

    add_body_para(doc, '2.2  Design Phase (Weeks 5-10): Based on the assessment findings, Provider shall develop a '
                  'detailed system architecture design, data migration plan, integration specifications, and custom '
                  'module requirements documentation.')

    add_body_para(doc, '2.3  Development Phase (Weeks 11-26): Provider shall configure the CloudHealth Enterprise Suite, '
                  'develop custom modules for Client\'s specialized workflows (including oncology, cardiology, and '
                  'rehabilitation departments), and build data migration scripts.')

    add_body_para(doc, '2.4  Testing Phase (Weeks 27-32): Provider shall perform comprehensive testing including unit '
                  'testing, integration testing, user acceptance testing, and HIPAA compliance validation.')

    add_body_para(doc, '2.5  Deployment Phase (Weeks 33-38): Provider shall execute a phased deployment across Client\'s '
                  'facilities, beginning with the flagship Seattle Regional Medical Center.')

    add_body_para(doc, '2.6  Training and Support Phase (Weeks 39-48): Provider shall deliver role-based training to '
                  'approximately 3,200 end users and provide on-site support during the transition period.')

    add_body_para(doc, '3. DELIVERABLES', bold=True)
    add_body_para(doc, '3.1  Technical Assessment Report')
    add_body_para(doc, '3.2  System Architecture Design Document')
    add_body_para(doc, '3.3  Data Migration Plan and Scripts')
    add_body_para(doc, '3.4  Configured CloudHealth Enterprise Suite')
    add_body_para(doc, '3.5  Custom Module Source Code and Documentation')
    add_body_para(doc, '3.6  Testing Reports (unit, integration, UAT, compliance)')
    add_body_para(doc, '3.7  Training Materials and User Guides')
    add_body_para(doc, '3.8  Post-Deployment Support Documentation')

    add_body_para(doc, '4. TIMELINE AND MILESTONES', bold=True)
    add_body_para(doc, 'Total project duration: Forty-eight (48) weeks from SOW execution date.')
    add_body_para(doc, 'Milestone 1 — Assessment Complete: Week 4')
    add_body_para(doc, 'Milestone 2 — Design Approved: Week 10')
    add_body_para(doc, 'Milestone 3 — Development Complete: Week 26')
    add_body_para(doc, 'Milestone 4 — UAT Sign-off: Week 32')
    add_body_para(doc, 'Milestone 5 — Full Deployment: Week 38')
    add_body_para(doc, 'Milestone 6 — Project Closure: Week 48')

    add_body_para(doc, '5. FEES', bold=True)
    add_body_para(doc, 'Total estimated project fee: $4,750,000.00 (Four Million Seven Hundred Fifty Thousand Dollars).')
    add_body_para(doc, 'Payment schedule: Monthly invoicing based on actual hours and expenses, subject to milestone '
                  'verification by the joint steering committee.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
