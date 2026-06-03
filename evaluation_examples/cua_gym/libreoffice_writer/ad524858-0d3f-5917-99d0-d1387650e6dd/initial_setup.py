"""
Initial Setup: 2-page draft proposal document without any watermark/text box
Task ID: writer_obj_059
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'draft_proposal'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set page margins (standard)
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ----- PAGE 1 -----
    # Title
    title = doc.add_heading('Project Proposal: Digital Transformation Initiative', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(12)

    # Subtitle / metadata
    meta = doc.add_paragraph()
    meta.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run('Prepared by: Strategic Planning Department')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    meta2 = doc.add_paragraph()
    meta2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = meta2.add_run('Date: March 15, 2025  |  Version: 0.9 (Under Review)')
    run2.font.name = 'Calibri'
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()  # spacer

    # Executive Summary heading
    doc.add_heading('1. Executive Summary', level=1)

    p1 = doc.add_paragraph(
        'This proposal outlines a comprehensive digital transformation strategy for Meridian Solutions Inc. '
        'to modernize core business operations, enhance customer experience, and drive sustainable growth '
        'over the next three fiscal years. The initiative encompasses cloud migration, process automation, '
        'and workforce upskilling programmes aligned with our 2025–2028 corporate strategy.'
    )
    p1.paragraph_format.space_after = Pt(8)

    p2 = doc.add_paragraph(
        'Total projected investment is estimated at $4.2 million, with an expected ROI of 220% within '
        'three years. Key performance indicators include a 35% reduction in operational costs, a 50% '
        'improvement in customer onboarding time, and a 40% increase in data-driven decision-making capacity.'
    )
    p2.paragraph_format.space_after = Pt(8)

    # Objectives heading
    doc.add_heading('2. Objectives', level=1)

    objectives = [
        'Migrate 85% of on-premise infrastructure to cloud-based solutions by Q4 2025.',
        'Implement end-to-end process automation across Finance, HR, and Supply Chain departments.',
        'Deploy an enterprise-wide data analytics platform to support real-time reporting.',
        'Achieve ISO 27001 certification for information security management by mid-2026.',
        'Upskill 1,200 employees through structured digital literacy and tools training programmes.',
    ]
    for obj in objectives:
        p = doc.add_paragraph(obj, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()  # spacer

    # Scope section
    doc.add_heading('3. Scope of Work', level=1)

    p3 = doc.add_paragraph(
        'The scope of this engagement covers three primary workstreams: Infrastructure Modernisation, '
        'Application Portfolio Rationalisation, and Change Management. Each workstream has been scoped '
        'with defined deliverables, timelines, and accountable stakeholders.'
    )
    p3.paragraph_format.space_after = Pt(8)

    # Table of workstreams
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    headers = ['Workstream', 'Lead', 'Target Completion']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        run_h = cell.paragraphs[0].add_run(h)
        run_h.bold = True
        run_h.font.size = Pt(10)
        run_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)

    data_rows = [
        ['Infrastructure Modernisation', 'Jessica Park', 'Q4 2025'],
        ['Application Portfolio Rationalisation', 'David Okonkwo', 'Q2 2026'],
        ['Change Management & Training', 'Layla Fernandez', 'Q3 2026'],
    ]
    for i, row_data in enumerate(data_rows, 1):
        for j, val in enumerate(row_data):
            cell = table.cell(i, j)
            run_d = cell.paragraphs[0].add_run(val)
            run_d.font.size = Pt(10)

    doc.add_paragraph()  # spacer after table

    # ----- PAGE 2 -----
    doc.add_page_break()

    doc.add_heading('4. Financial Overview', level=1)

    p4 = doc.add_paragraph(
        'The proposed budget has been structured across three phases, ensuring that capital expenditure '
        'is staged to align with business readiness and organisational capacity. Phase 1 (2025) accounts '
        'for 45% of total investment, primarily covering infrastructure and licensing costs.'
    )
    p4.paragraph_format.space_after = Pt(8)

    # Budget table
    budget_table = doc.add_table(rows=5, cols=3)
    budget_table.style = 'Table Grid'

    bheaders = ['Phase', 'Period', 'Estimated Cost']
    for j, h in enumerate(bheaders):
        cell = budget_table.cell(0, j)
        run_bh = cell.paragraphs[0].add_run(h)
        run_bh.bold = True
        run_bh.font.size = Pt(10)

    budget_data = [
        ['Phase 1 – Foundation', '2025', '$1,890,000'],
        ['Phase 2 – Expansion', '2026', '$1,470,000'],
        ['Phase 3 – Optimisation', '2027', '$840,000'],
        ['Total', '', '$4,200,000'],
    ]
    for i, row_data in enumerate(budget_data, 1):
        for j, val in enumerate(row_data):
            cell = budget_table.cell(i, j)
            run_bd = cell.paragraphs[0].add_run(val)
            run_bd.font.size = Pt(10)
            if i == 4:
                run_bd.bold = True

    doc.add_paragraph()

    doc.add_heading('5. Risk Assessment', level=1)

    risks = [
        ('Vendor Lock-in', 'Medium', 'Multi-cloud architecture with open standards adoption.'),
        ('Change Resistance', 'High', 'Executive sponsorship and structured change management programme.'),
        ('Budget Overrun', 'Low', 'Stage-gate approvals and contingency reserve of 10%.'),
        ('Data Security Breach', 'Medium', 'ISO 27001 roadmap and third-party security audits.'),
    ]

    risk_table = doc.add_table(rows=len(risks) + 1, cols=3)
    risk_table.style = 'Table Grid'

    rh = ['Risk', 'Likelihood', 'Mitigation Strategy']
    for j, h in enumerate(rh):
        cell = risk_table.cell(0, j)
        run_rh = cell.paragraphs[0].add_run(h)
        run_rh.bold = True
        run_rh.font.size = Pt(10)

    for i, (risk, likelihood, mitigation) in enumerate(risks, 1):
        risk_table.cell(i, 0).paragraphs[0].add_run(risk).font.size = Pt(10)
        risk_table.cell(i, 1).paragraphs[0].add_run(likelihood).font.size = Pt(10)
        risk_table.cell(i, 2).paragraphs[0].add_run(mitigation).font.size = Pt(10)

    doc.add_paragraph()

    doc.add_heading('6. Conclusion', level=1)

    p5 = doc.add_paragraph(
        'The Digital Transformation Initiative represents a strategic imperative for Meridian Solutions Inc. '
        'to remain competitive in an increasingly digital-first marketplace. With strong leadership commitment, '
        'phased investment, and a robust change management framework, the organisation is well-positioned to '
        'realise the full benefits of this transformation.'
    )
    p5.paragraph_format.space_after = Pt(8)

    p6 = doc.add_paragraph(
        'We seek approval from the Executive Committee to proceed with Phase 1 activities commencing '
        'April 2025. A detailed project charter and governance framework will be presented within '
        '30 days of approval.'
    )
    p6.paragraph_format.space_after = Pt(8)

    # Approval signature block
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(24)
    run_sig = sig.add_run('Submitted by: ___________________________   Date: _____________')
    run_sig.font.name = 'Calibri'
    run_sig.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
