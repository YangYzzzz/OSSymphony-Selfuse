"""
Initial Setup: Contacts list with phone numbers in (XXX) XXX-XXXX format
Task ID: writer_edit_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'  # VM path — task file is on Desktop
TASK_ID = 'contacts_list'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading('Contact Directory', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('Internal Team Contacts — Updated Q1 2025')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.size = Pt(11)

    doc.add_paragraph('')  # blank line

    # Contact data: name, title, email, phone in (XXX) XXX-XXXX format
    contacts = [
        {
            'name': 'Sarah Chen',
            'title': 'Senior Software Engineer',
            'email': 'sarah.chen@company.com',
            'phone': '(555) 100-2000',
        },
        {
            'name': 'Marcus Johnson',
            'title': 'Marketing Manager',
            'email': 'marcus.johnson@company.com',
            'phone': '(212) 555-0199',
        },
        {
            'name': 'Priya Patel',
            'title': 'Product Designer',
            'email': 'priya.patel@company.com',
            'phone': '(310) 444-7890',
        },
        {
            'name': 'Derek Thompson',
            'title': 'Data Analyst',
            'email': 'derek.thompson@company.com',
            'phone': '(408) 222-3344',
        },
        {
            'name': 'Caitlin Rivera',
            'title': 'HR Business Partner',
            'email': 'caitlin.rivera@company.com',
            'phone': '(617) 888-1122',
        },
        {
            'name': 'James Nakamura',
            'title': 'Finance Director',
            'email': 'james.nakamura@company.com',
            'phone': '(702) 333-9876',
        },
        {
            'name': 'Aisha Williams',
            'title': 'Operations Coordinator',
            'email': 'aisha.williams@company.com',
            'phone': '(503) 777-4455',
        },
        {
            'name': 'Roberto Martinez',
            'title': 'Sales Executive',
            'email': 'roberto.martinez@company.com',
            'phone': '(818) 666-5544',
        },
    ]

    # Add each contact as a formatted block
    for i, contact in enumerate(contacts, 1):
        # Contact heading: numbered name
        heading = doc.add_heading(f'{i}. {contact["name"]}', level=2)

        # Title / Position
        p_title = doc.add_paragraph()
        run_label = p_title.add_run('Title: ')
        run_label.font.bold = True
        run_label.font.size = Pt(11)
        run_value = p_title.add_run(contact['title'])
        run_value.font.size = Pt(11)

        # Email
        p_email = doc.add_paragraph()
        run_label = p_email.add_run('Email: ')
        run_label.font.bold = True
        run_label.font.size = Pt(11)
        run_value = p_email.add_run(contact['email'])
        run_value.font.size = Pt(11)

        # Phone — in (XXX) XXX-XXXX format (NOT yet converted)
        p_phone = doc.add_paragraph()
        run_label = p_phone.add_run('Phone: ')
        run_label.font.bold = True
        run_label.font.size = Pt(11)
        run_value = p_phone.add_run(contact['phone'])
        run_value.font.size = Pt(11)

        # Blank line between contacts (except last)
        if i < len(contacts):
            doc.add_paragraph('')

    # Footer note
    doc.add_paragraph('')
    note = doc.add_paragraph('Note: For urgent matters, please use the company emergency line at ext. 9000.')
    note.runs[0].font.italic = True
    note.runs[0].font.size = Pt(10)
    note.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
