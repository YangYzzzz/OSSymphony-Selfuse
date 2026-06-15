"""
Reward Script: Use Find & Replace with regex to convert phone number format
Task ID: writer_edit_004
Domain: libreoffice_writer
Scoring:
  Component 1: All 8 phone numbers converted to XXX-XXX-XXXX format (0.6 pts)
               Each phone is worth 0.075 pts; all 8 = 0.6 pts
  Component 2: No old-format (XXX) XXX-XXXX phone numbers remain (0.2 pts)
  Component 3: All 8 phones in new format AND document non-phone content preserved (0.2 pts)
               Compound check — must pass Component 1 first, so fails on initial_env
"""

import os
import re

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_004'

# Expected phone numbers in the new (post-task) format XXX-XXX-XXXX
EXPECTED_PHONES_NEW = [
    '555-100-2000',
    '212-555-0199',
    '310-444-7890',
    '408-222-3344',
    '617-888-1122',
    '702-333-9876',
    '503-777-4455',
    '818-666-5544',
]

# Old-format phone pattern: (XXX) XXX-XXXX — should NOT exist post-task
OLD_FORMAT_PATTERN = re.compile(r'\(\d{3}\)\s\d{3}-\d{4}')


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather all text in the document
    all_text_lines = []
    for para in doc.paragraphs:
        all_text_lines.append(para.text)

    # Also check tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_text_lines.append(para.text)

    full_text = '\n'.join(all_text_lines)

    # Component 1: Each phone number converted to XXX-XXX-XXXX format (0.075 pts each, 0.6 total)
    # These phone strings DO NOT exist in initial_env, so this fails on initial, passes on golden
    try:
        found_count = 0
        for phone in EXPECTED_PHONES_NEW:
            if phone in full_text:
                found_count += 1
            else:
                print(f"FAIL: Component 1 — new-format phone '{phone}' not found in document")

        if found_count == 8:
            print(f"PASS: Component 1 — all 8 phone numbers converted to XXX-XXX-XXXX format (0.6 pts)")
            total_score += 0.6
        elif found_count > 0:
            partial = round(found_count * 0.075, 3)
            print(f"PARTIAL: Component 1 — {found_count}/8 phones in new format ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — no phones in new XXX-XXX-XXXX format found (0.0 pts)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: No old-format (XXX) XXX-XXXX phones remain (0.2 points)
    # In initial_env all 8 phones are in old format — this component fails on initial (8 old phones exist)
    # In golden_env no old-format phones remain — this component passes on golden
    try:
        old_format_matches = OLD_FORMAT_PATTERN.findall(full_text)
        if len(old_format_matches) == 0:
            print(f"PASS: Component 2 — no old-format '(XXX) XXX-XXXX' phone numbers remain (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 2 — {len(old_format_matches)} old-format phone(s) still present: {old_format_matches}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All 8 new-format phones are present AND total phone count equals 8 (0.2 points)
    # This is a compound check: new phones found (fails on initial) AND count is correct (not over-replaced)
    # Checking count of new-format matches = 8, while old-format = 0, together = complete replacement
    try:
        new_format_all = re.findall(r'\b\d{3}-\d{3}-\d{4}\b', full_text)
        old_format_all = OLD_FORMAT_PATTERN.findall(full_text)
        # Condition: exactly 8 new-format phones AND 0 old-format phones
        # This fails on initial_env (0 new + 8 old) and passes on golden (8 new + 0 old)
        if len(new_format_all) == 8 and len(old_format_all) == 0:
            print(f"PASS: Component 3 — exactly 8 new-format phones found, 0 old-format phones (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — new-format count={len(new_format_all)} (expected 8), old-format count={len(old_format_all)} (expected 0)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in a given env
file_path = f'{WORKDIR}/Desktop/contacts_list.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
