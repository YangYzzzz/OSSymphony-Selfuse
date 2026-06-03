"""
Initial Setup: Create a master document with three subdocuments (no TOC)
Task ID: writer_af_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
import zipfile
from xml.etree.ElementTree import Element, SubElement, tostring

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_af_031'

CHAPTER1 = f'{WORKDIR}/chapter1.docx'
CHAPTER2 = f'{WORKDIR}/chapter2.docx'
CHAPTER3 = f'{WORKDIR}/chapter3.docx'
MASTER_DOC = f'{WORKDIR}/Complete_Thesis.odm'

NS = {
    'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0',
    'style': 'urn:oasis:names:tc:opendocument:xmlns:style:1.0',
    'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
    'table': 'urn:oasis:names:tc:opendocument:xmlns:table:1.0',
    'draw': 'urn:oasis:names:tc:opendocument:xmlns:drawing:1.0',
    'fo': 'urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0',
    'xlink': 'http://www.w3.org/1999/xlink',
    'svg': 'urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0',
    'manifest': 'urn:oasis:names:tc:opendocument:xmlns:manifest:1.0',
}


def launch_gui(command, delay_sec=1.0):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_chapter1():
    doc = Document()
    doc.add_heading('Introduction to Machine Learning', level=1)
    doc.add_paragraph(
        'Machine learning has become one of the most transformative technologies '
        'of the twenty-first century. From autonomous vehicles to medical diagnostics, '
        'its applications span virtually every industry. This chapter provides a '
        'foundational overview of the key concepts, history, and current landscape '
        'of machine learning research and applications.'
    )
    doc.add_heading('Historical Background', level=2)
    doc.add_paragraph(
        'The origins of machine learning can be traced back to the 1950s, when '
        'Arthur Samuel developed a program that could learn to play checkers. '
        'Since then, the field has undergone several waves of enthusiasm and '
        'disillusionment, often referred to as "AI winters." The resurgence of '
        'interest in the 2010s, driven by deep learning breakthroughs, has '
        'established machine learning as a central pillar of modern computing.'
    )
    doc.add_paragraph(
        'Key milestones include the development of the perceptron by Frank '
        'Rosenblatt in 1957, the backpropagation algorithm popularized in the '
        '1980s, and the introduction of convolutional neural networks for image '
        'recognition in the 1990s by Yann LeCun and colleagues.'
    )
    doc.add_heading('Types of Machine Learning', level=2)
    doc.add_paragraph(
        'Machine learning algorithms are broadly categorized into three types: '
        'supervised learning, unsupervised learning, and reinforcement learning. '
        'Supervised learning involves training on labeled data, where the algorithm '
        'learns to map inputs to known outputs. Common applications include image '
        'classification, spam detection, and medical diagnosis.'
    )
    doc.add_paragraph(
        'Unsupervised learning, on the other hand, deals with unlabeled data. '
        'The algorithm must discover hidden patterns or structures within the data. '
        'Clustering, dimensionality reduction, and anomaly detection are typical '
        'unsupervised learning tasks.'
    )
    doc.add_paragraph(
        'Reinforcement learning involves an agent that learns optimal behavior '
        'through interaction with an environment, receiving rewards or penalties '
        'based on its actions. This approach has achieved remarkable success in '
        'game playing, robotics, and resource management.'
    )
    doc.add_heading('Current Challenges', level=2)
    doc.add_paragraph(
        'Despite significant progress, machine learning faces several challenges. '
        'Data quality and availability remain persistent issues, particularly in '
        'specialized domains such as healthcare and finance. Model interpretability '
        'is another concern, as complex models often function as "black boxes" that '
        'are difficult to understand or debug.'
    )
    doc.save(CHAPTER1)
    print(f'Created: {CHAPTER1}')


def create_chapter2():
    doc = Document()
    doc.add_heading('Neural Network Architectures', level=1)
    doc.add_paragraph(
        'Neural networks form the backbone of modern deep learning systems. '
        'This chapter examines the most important architectural innovations that '
        'have driven the field forward, from simple feedforward networks to '
        'sophisticated transformer-based models.'
    )
    doc.add_heading('Feedforward Neural Networks', level=2)
    doc.add_paragraph(
        'The feedforward neural network, also known as a multilayer perceptron (MLP), '
        'is the simplest form of artificial neural network. Information flows in one '
        'direction, from input nodes through hidden layers to output nodes. Despite '
        'their simplicity, feedforward networks can approximate any continuous '
        'function given sufficient width and depth.'
    )
    doc.add_paragraph(
        'Training feedforward networks relies on the backpropagation algorithm '
        'combined with gradient descent optimization. Modern variants use '
        'techniques such as batch normalization, dropout regularization, and '
        'adaptive learning rate schedules to improve convergence and generalization.'
    )
    doc.add_heading('Convolutional Neural Networks', level=2)
    doc.add_paragraph(
        'Convolutional Neural Networks (CNNs) were designed specifically for '
        'processing grid-structured data such as images. They use convolutional '
        'layers that apply learnable filters to detect local features, followed '
        'by pooling layers that reduce spatial dimensions. Landmark architectures '
        'include LeNet-5, AlexNet, VGGNet, ResNet, and EfficientNet.'
    )
    doc.add_paragraph(
        'The introduction of residual connections by He et al. in 2015 was a '
        'breakthrough that enabled training of much deeper networks. ResNet '
        'demonstrated that networks with over 100 layers could be trained '
        'effectively, achieving human-level performance on image classification.'
    )
    doc.add_heading('Recurrent Neural Networks and Transformers', level=2)
    doc.add_paragraph(
        'Recurrent Neural Networks (RNNs) and their variants, such as Long '
        'Short-Term Memory (LSTM) networks and Gated Recurrent Units (GRUs), '
        'were the dominant architecture for sequential data processing until '
        'the advent of transformers.'
    )
    doc.add_paragraph(
        'The transformer architecture, introduced by Vaswani et al. in 2017, '
        'replaced recurrence with self-attention mechanisms, enabling parallel '
        'processing of sequences and capturing long-range dependencies more '
        'effectively. Transformers have become the foundation for large language '
        'models such as GPT, BERT, and their successors.'
    )
    doc.add_heading('Generative Adversarial Networks', level=2)
    doc.add_paragraph(
        'Generative Adversarial Networks (GANs), introduced by Goodfellow et '
        'al. in 2014, consist of two competing neural networks: a generator '
        'that creates synthetic data and a discriminator that distinguishes '
        'real from generated samples. GANs have achieved impressive results in '
        'image synthesis, style transfer, and data augmentation.'
    )
    doc.save(CHAPTER2)
    print(f'Created: {CHAPTER2}')


def create_chapter3():
    doc = Document()
    doc.add_heading('Applications and Future Directions', level=1)
    doc.add_paragraph(
        'The practical applications of machine learning continue to expand '
        'across industries. This chapter surveys key application domains and '
        'discusses emerging trends that are likely to shape the future of '
        'the field.'
    )
    doc.add_heading('Healthcare Applications', level=2)
    doc.add_paragraph(
        'Machine learning is transforming healthcare through improved '
        'diagnostics, drug discovery, and personalized medicine. Deep learning '
        'models have demonstrated expert-level performance in detecting '
        'diseases from medical images, including diabetic retinopathy, skin '
        'cancer, and pneumonia from chest X-rays.'
    )
    doc.add_paragraph(
        'In drug discovery, machine learning accelerates the identification '
        'of promising molecular compounds by predicting binding affinities '
        'and toxicity profiles. AlphaFold, developed by DeepMind, revolutionized '
        'structural biology by accurately predicting protein structures from '
        'amino acid sequences.'
    )
    doc.add_heading('Autonomous Systems', level=2)
    doc.add_paragraph(
        'Self-driving vehicles represent one of the most ambitious applications '
        'of machine learning. Companies such as Waymo, Tesla, and Cruise are '
        'developing systems that combine computer vision, sensor fusion, and '
        'reinforcement learning to navigate complex driving scenarios.'
    )
    doc.add_paragraph(
        'Beyond transportation, autonomous systems are being deployed in '
        'agriculture for precision farming, in manufacturing for quality '
        'inspection, and in logistics for warehouse automation. These systems '
        'rely heavily on real-time perception and decision-making capabilities '
        'powered by neural networks.'
    )
    doc.add_heading('Natural Language Processing', level=2)
    doc.add_paragraph(
        'Large language models have dramatically advanced the state of natural '
        'language processing. Applications include machine translation, text '
        'summarization, question answering, and conversational AI. The scale '
        'of these models continues to grow, with recent systems containing '
        'hundreds of billions of parameters.'
    )
    doc.add_heading('Ethical Considerations and Future Outlook', level=2)
    doc.add_paragraph(
        'As machine learning systems become more pervasive, ethical concerns '
        'have moved to the forefront. Issues of algorithmic bias, privacy, '
        'transparency, and accountability demand careful attention from '
        'researchers, developers, and policymakers. The development of '
        'responsible AI frameworks and governance structures will be crucial '
        'for ensuring that these powerful technologies benefit society broadly.'
    )
    doc.add_paragraph(
        'Looking ahead, key research directions include more efficient training '
        'methods, improved model interpretability, multimodal learning systems, '
        'and the development of artificial general intelligence. The convergence '
        'of machine learning with other fields such as neuroscience, quantum '
        'computing, and robotics promises to unlock new frontiers of innovation.'
    )
    doc.save(CHAPTER3)
    print(f'Created: {CHAPTER3}')


def make_content_xml_no_toc():
    root = Element('office:document-content')
    root.set('xmlns:office', NS['office'])
    root.set('xmlns:style', NS['style'])
    root.set('xmlns:text', NS['text'])
    root.set('xmlns:table', NS['table'])
    root.set('xmlns:draw', NS['draw'])
    root.set('xmlns:fo', NS['fo'])
    root.set('xmlns:xlink', NS['xlink'])
    root.set('xmlns:svg', NS['svg'])
    root.set('office:version', '1.2')

    SubElement(root, 'office:automatic-styles')

    body = SubElement(root, 'office:body')
    global_doc = SubElement(body, 'office:global-document')

    for i, fname in enumerate(['chapter1.docx', 'chapter2.docx', 'chapter3.docx'], 1):
        section = SubElement(global_doc, 'text:section')
        section.set('text:name', f'Chapter {i}')
        section.set('text:protected', 'false')
        source = SubElement(section, 'text:section-source')
        source.set('xlink:href', fname)
        source.set('xlink:type', 'simple')
        source.set('text:section-name', '')
        source.set('text:filter-name', 'MS Word 2007 XML')

    return tostring(root, encoding='unicode', xml_declaration=True)


def make_styles_xml():
    root = Element('office:document-styles')
    root.set('xmlns:office', NS['office'])
    root.set('xmlns:style', NS['style'])
    root.set('xmlns:text', NS['text'])
    root.set('xmlns:fo', NS['fo'])
    root.set('office:version', '1.2')

    styles = SubElement(root, 'office:styles')
    dp = SubElement(styles, 'style:default-style')
    dp.set('style:family', 'paragraph')
    pp = SubElement(dp, 'style:paragraph-properties')
    pp.set('fo:margin-top', '0in')
    pp.set('fo:margin-bottom', '0.08in')
    tp = SubElement(dp, 'style:text-properties')
    tp.set('fo:font-size', '12pt')
    tp.set('style:font-name', 'Liberation Serif')

    auto_styles = SubElement(root, 'office:automatic-styles')
    master_styles = SubElement(root, 'office:master-styles')
    mp = SubElement(master_styles, 'style:master-page')
    mp.set('style:name', 'Standard')
    mp.set('style:page-layout-name', 'pm1')
    pl = SubElement(auto_styles, 'style:page-layout')
    pl.set('style:name', 'pm1')
    plp = SubElement(pl, 'style:page-layout-properties')
    plp.set('fo:page-width', '8.5in')
    plp.set('fo:page-height', '11in')
    plp.set('fo:margin-top', '1in')
    plp.set('fo:margin-bottom', '1in')
    plp.set('fo:margin-left', '1in')
    plp.set('fo:margin-right', '1in')

    return tostring(root, encoding='unicode', xml_declaration=True)


def make_meta_xml():
    root = Element('office:document-meta')
    root.set('xmlns:office', NS['office'])
    root.set('office:version', '1.2')
    SubElement(root, 'office:meta')
    return tostring(root, encoding='unicode', xml_declaration=True)


def make_manifest_xml():
    root = Element('manifest:manifest')
    root.set('xmlns:manifest', NS['manifest'])
    root.set('manifest:version', '1.2')

    e1 = SubElement(root, 'manifest:file-entry')
    e1.set('manifest:media-type', 'application/vnd.oasis.opendocument.text-master')
    e1.set('manifest:full-path', '/')

    e2 = SubElement(root, 'manifest:file-entry')
    e2.set('manifest:media-type', 'text/xml')
    e2.set('manifest:full-path', 'content.xml')

    e3 = SubElement(root, 'manifest:file-entry')
    e3.set('manifest:media-type', 'text/xml')
    e3.set('manifest:full-path', 'styles.xml')

    e4 = SubElement(root, 'manifest:file-entry')
    e4.set('manifest:media-type', 'text/xml')
    e4.set('manifest:full-path', 'meta.xml')

    return tostring(root, encoding='unicode', xml_declaration=True)


def create_master_document():
    with zipfile.ZipFile(MASTER_DOC, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('mimetype', 'application/vnd.oasis.opendocument.text-master')
        zf.writestr('content.xml', make_content_xml_no_toc())
        zf.writestr('styles.xml', make_styles_xml())
        zf.writestr('meta.xml', make_meta_xml())
        zf.writestr('META-INF/manifest.xml', make_manifest_xml())
    print(f'Master document created: {MASTER_DOC}')


def main():
    create_chapter1()
    create_chapter2()
    create_chapter3()
    create_master_document()
    launch_gui(f'libreoffice --writer "{MASTER_DOC}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with Complete_Thesis.odm')


main()
