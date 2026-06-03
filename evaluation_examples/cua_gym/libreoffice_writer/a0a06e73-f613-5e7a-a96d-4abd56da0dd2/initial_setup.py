"""
Initial Setup: Create five chapter subdocuments for a thesis master document.
Task ID: writer_acad_053
Domain: libreoffice_writer

Creates five .docx files (chapter1-5) with realistic thesis chapter content,
then opens a new empty Writer document for the agent to use as the master document.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_053'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)

def create_chapter(filename, title, content_paragraphs):
    """Create a chapter document with realistic thesis content."""
    doc = Document()

    # Chapter heading
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Chapter content
    for para_text in content_paragraphs:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    doc.save(os.path.join(WORKDIR, filename))
    print(f'Created: {filename}')

def create_initial():
    # Chapter 1: Introduction
    create_chapter('chapter1.docx', 'Chapter 1: Introduction', [
        'The rapid advancement of artificial intelligence and machine learning has fundamentally '
        'transformed how researchers approach complex data analysis problems across multiple '
        'scientific disciplines. This thesis investigates the application of deep neural networks '
        'to climate prediction models, aiming to improve forecast accuracy beyond traditional '
        'numerical weather prediction methods.',
        'Climate modeling has historically relied on physics-based simulations that solve '
        'partial differential equations governing atmospheric dynamics. While these models '
        'have achieved remarkable success, they require enormous computational resources '
        'and often struggle with sub-grid-scale processes that cannot be explicitly resolved '
        'at current resolutions (Palmer & Stevens, 2019).',
        'Recent studies have demonstrated that data-driven approaches, particularly those '
        'leveraging convolutional neural networks and transformer architectures, can capture '
        'spatiotemporal patterns in meteorological data that complement physics-based models '
        '(Ravuri et al., 2021; Bi et al., 2023). This thesis builds upon these findings '
        'by proposing a hybrid framework that integrates learned representations with '
        'established physical constraints.',
        'The primary research questions addressed in this thesis are: (1) Can deep learning '
        'models improve medium-range weather forecasts when combined with numerical weather '
        'prediction output? (2) What architectural choices are most effective for capturing '
        'multi-scale atmospheric phenomena? (3) How can physical conservation laws be '
        'incorporated into neural network training to ensure physically consistent predictions?',
    ])

    # Chapter 2: Literature Review
    create_chapter('chapter2.docx', 'Chapter 2: Literature Review', [
        'This chapter provides a comprehensive review of the literature spanning three key '
        'areas: traditional numerical weather prediction, machine learning for geoscientific '
        'applications, and hybrid physics-informed neural network approaches.',
        'Numerical weather prediction (NWP) has evolved significantly since the pioneering '
        'work of Richardson (1922) and the first successful computer-based forecast by '
        'Charney, Fjortoft, and von Neumann (1950). Modern operational systems such as the '
        'European Centre for Medium-Range Weather Forecasts (ECMWF) Integrated Forecasting '
        'System achieve skillful predictions up to 10 days ahead for large-scale patterns, '
        'with resolution reaching approximately 9 km globally (Bauer et al., 2015).',
        'The application of neural networks to weather prediction dates back to the work of '
        'Hall et al. (1999), who demonstrated that multilayer perceptrons could capture '
        'nonlinear relationships in meteorological variables. The field experienced rapid '
        'growth with the advent of deep learning, as convolutional architectures proved '
        'particularly effective at processing gridded atmospheric data (Dueben & Bauer, 2018).',
        'Physics-informed neural networks (PINNs), introduced by Raissi et al. (2019), '
        'represent a paradigm shift in scientific machine learning. By encoding differential '
        'equations directly into the loss function, PINNs can produce solutions that respect '
        'fundamental physical laws while leveraging the flexibility of neural network function '
        'approximation. Extensions to atmospheric science have shown promise in ensuring '
        'conservation of mass, energy, and momentum (Beucler et al., 2021).',
        'Transfer learning and foundation models have emerged as powerful tools for climate '
        'science. Nguyen et al. (2023) demonstrated that large pretrained vision transformers '
        'can be fine-tuned for regional weather prediction with limited labeled data, '
        'significantly reducing the computational cost of developing specialized forecasting '
        'systems for under-resourced regions.',
    ])

    # Chapter 3: Methodology
    create_chapter('chapter3.docx', 'Chapter 3: Methodology', [
        'This chapter describes the methodological framework developed for integrating deep '
        'learning with numerical weather prediction. The approach consists of three main '
        'components: data preprocessing and feature engineering, model architecture design, '
        'and physics-constrained training procedures.',
        'The training dataset comprises ERA5 reanalysis data from 1979 to 2022, provided by '
        'the Copernicus Climate Change Service at 0.25-degree spatial resolution with hourly '
        'temporal resolution. Variables include geopotential height, temperature, specific '
        'humidity, and wind components at 13 pressure levels, along with surface variables '
        'such as mean sea level pressure, 2-meter temperature, and total precipitation.',
        'The proposed architecture, termed ClimateNet-Hybrid, employs a U-Net backbone with '
        'axial attention mechanisms to capture both local and global spatial dependencies. '
        'The encoder processes input fields at multiple resolutions through a series of '
        'convolutional blocks with skip connections, while the decoder reconstructs forecast '
        'fields at the target resolution. Temporal dynamics are modeled through an autoregressive '
        'rollout scheme with 6-hour time steps.',
        'To enforce physical consistency, we augment the standard mean squared error loss with '
        'penalty terms derived from the continuity equation, thermodynamic energy equation, '
        'and hydrostatic balance. These soft constraints are weighted using a multi-task '
        'learning framework with learned task weights following Kendall et al. (2018). '
        'The total loss function is formulated as L = L_mse + lambda_1 * L_mass + '
        'lambda_2 * L_energy + lambda_3 * L_hydrostatic.',
        'Model training is performed on 8 NVIDIA A100 GPUs using distributed data-parallel '
        'training with a batch size of 32 and the AdamW optimizer with cosine learning rate '
        'scheduling. The initial learning rate is set to 3e-4 with a warmup period of 1000 '
        'steps. Training proceeds for 100 epochs with early stopping based on validation RMSE.',
    ])

    # Chapter 4: Results and Discussion
    create_chapter('chapter4.docx', 'Chapter 4: Results and Discussion', [
        'This chapter presents the experimental results of ClimateNet-Hybrid evaluated against '
        'operational NWP baselines and state-of-the-art machine learning approaches. We report '
        'results for both deterministic and probabilistic forecast metrics across multiple '
        'lead times and atmospheric variables.',
        'For 500 hPa geopotential height, ClimateNet-Hybrid achieves a root mean squared error '
        'of 87.3 m at 72-hour lead time, compared to 94.1 m for the operational ECMWF HRES '
        'and 91.6 m for Pangu-Weather (Bi et al., 2023). The improvement is statistically '
        'significant at the 95% confidence level using a paired bootstrap test (Hamill, 1999) '
        'with 10,000 resamples over the 2020-2022 test period.',
        'Surface temperature forecasts show a 12.4% reduction in bias compared to the NWP '
        'baseline, with the largest improvements observed in coastal regions where sub-grid '
        'processes related to land-sea contrasts pose challenges for coarse-resolution numerical '
        'models. The physics constraints prove particularly beneficial for these regions, '
        'reducing unphysical temperature gradients by 34.7% compared to the unconstrained '
        'deep learning baseline.',
        'Analysis of the learned feature representations reveals that the model develops '
        'internally consistent representations of large-scale atmospheric circulation patterns. '
        'Singular value decomposition of the bottleneck layer activations shows strong '
        'correspondence with known teleconnection patterns including the North Atlantic '
        'Oscillation (NAO) and El Nino-Southern Oscillation (ENSO), suggesting that the '
        'network discovers physically meaningful modes of variability without explicit supervision.',
        'A key limitation emerges for extreme weather events. While ClimateNet-Hybrid improves '
        'average forecast skill, it tends to underpredict the intensity of extreme precipitation '
        'events by 15-20%, consistent with the regression-to-the-mean behavior commonly '
        'observed in deterministic deep learning models trained with squared error losses.',
    ])

    # Chapter 5: Conclusions and Future Work
    create_chapter('chapter5.docx', 'Chapter 5: Conclusions and Future Work', [
        'This thesis has presented ClimateNet-Hybrid, a physics-informed deep learning framework '
        'for medium-range weather prediction that demonstrates significant improvements over '
        'both traditional numerical weather prediction and purely data-driven approaches. '
        'The key contributions and findings are summarized below.',
        'First, we established that incorporating physical conservation laws as soft constraints '
        'during neural network training produces forecasts that are not only more accurate '
        'but also more physically consistent than unconstrained deep learning approaches. '
        'The multi-task learning framework for balancing multiple physics penalties proved '
        'essential for stable training and optimal forecast performance.',
        'Second, the proposed axial attention mechanism within the U-Net architecture '
        'effectively captures multi-scale atmospheric interactions, from mesoscale convective '
        'systems to planetary-scale Rossby waves. This architectural choice represents a '
        'practical compromise between the computational efficiency of convolutional networks '
        'and the global receptive field of full attention transformers.',
        'Third, our analysis demonstrates that physics-informed training not only improves '
        'quantitative metrics but also yields more interpretable learned representations that '
        'correspond to known modes of atmospheric variability.',
        'Future research directions include: (1) extending the framework to subseasonal-to-seasonal '
        'prediction timescales, where the role of slowly varying boundary conditions becomes '
        'paramount; (2) developing probabilistic extensions through ensemble or diffusion-based '
        'approaches to better capture forecast uncertainty; and (3) investigating the potential '
        'for transfer learning to regional downscaling applications, enabling high-resolution '
        'forecasts in data-sparse regions of the developing world.',
    ])

    print(f'All five chapter documents created in {WORKDIR}')

    # Launch a new empty Writer document for the master document workspace
    # Kill any existing LibreOffice instances first for idempotency
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    time.sleep(1)

    launch_gui('libreoffice --writer', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')

create_initial()
