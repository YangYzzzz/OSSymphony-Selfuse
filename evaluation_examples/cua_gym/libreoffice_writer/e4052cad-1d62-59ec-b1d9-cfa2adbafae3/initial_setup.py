"""
Initial Setup: Weekly Fitness Tracker document (pre-task state)
Task ID: writer_tbl_057
Domain: libreoffice_writer

Creates fitness_log.docx on the Desktop with heading and paragraph,
but NO table yet (the agent must create the table).
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_057'
OUTPUT = f'{WORKDIR}/fitness_log.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Add heading: "Weekly Fitness Tracker"
    doc.add_heading('Weekly Fitness Tracker', level=1)

    # Add introductory paragraph
    doc.add_paragraph('Use this table to track your daily exercise routine.')

    # Save — NO table at this point (agent must create it)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer with DISPLAY=:0
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
