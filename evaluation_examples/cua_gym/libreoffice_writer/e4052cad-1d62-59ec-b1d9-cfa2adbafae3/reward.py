"""
Reward Script: Create fitness tracking table in LibreOffice Writer document
Task ID: writer_tbl_057
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Table exists with exactly 8 rows and 4 columns
  Component 2 (0.30): Header row contains correct column names
  Component 3 (0.20): Day column (col 0, rows 1-7) contains Mon-Sun weekdays
  Component 4 (0.20): All cell borders are 0.75pt (sz=6) solid dark green (#006400)
Total: 1.0
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_057'
FILE_PATH = f'{WORKDIR}/fitness_log.docx'

EXPECTED_HEADERS = ['Day', 'Exercise', 'Duration (min)', 'Calories']
EXPECTED_DAYS = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
EXPECTED_BORDER_SZ = '6'       # 0.75pt = 6 in Word sz units (1pt = 8 units)
EXPECTED_BORDER_COLOR = '006400'  # dark green
EXPECTED_BORDER_VAL = 'single'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition gate: file must exist and be loadable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: at least one table must exist to proceed with table checks
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document (task requires a table to be created)")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    table = doc.tables[0]

    # Component 1: Table has exactly 8 rows and 4 columns (0.30 points)
    # This FAILS on initial (no table) and PASSES on golden (8x4 table)
    try:
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        if num_rows == 8 and num_cols == 4:
            print(f"PASS: Component 1 — Table has 8 rows x 4 columns (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 1 — Expected 8 rows x 4 cols, found {num_rows} rows x {num_cols} cols")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header row contains correct column names (0.30 points)
    # This FAILS on initial (no table) and PASSES on golden (correct headers)
    try:
        header_cells = [table.cell(0, ci).text.strip() for ci in range(min(4, len(table.columns)))]
        if header_cells == EXPECTED_HEADERS:
            print(f"PASS: Component 2 — Header row is correct: {header_cells} (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 2 — Expected headers {EXPECTED_HEADERS}, found {header_cells}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Day column (rows 1-7) contains Monday through Sunday (0.20 points)
    # This FAILS on initial (no table) and PASSES on golden (days filled in)
    try:
        day_cells = [table.cell(ri, 0).text.strip() for ri in range(1, min(8, len(table.rows)))]
        if day_cells == EXPECTED_DAYS:
            print(f"PASS: Component 3 — Day column filled correctly: {day_cells} (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — Expected days {EXPECTED_DAYS}, found {day_cells}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: All cell borders are 0.75pt (sz=6) solid dark green (#006400) (0.20 points)
    # This FAILS on initial (no table) and PASSES on golden (dark green borders)
    try:
        border_issues = []
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                tc_pr = cell._tc.find(qn('w:tcPr'))
                if tc_pr is None:
                    border_issues.append(f"Cell [{ri},{ci}]: No tcPr")
                    continue
                tc_borders = tc_pr.find(qn('w:tcBorders'))
                if tc_borders is None:
                    border_issues.append(f"Cell [{ri},{ci}]: No tcBorders")
                    continue
                for side in ['top', 'left', 'bottom', 'right']:
                    border_elem = tc_borders.find(qn(f'w:{side}'))
                    if border_elem is None:
                        border_issues.append(f"Cell [{ri},{ci}]: Missing '{side}' border")
                    else:
                        sz = border_elem.get(qn('w:sz'))
                        color = border_elem.get(qn('w:color'))
                        val = border_elem.get(qn('w:val'))
                        if sz != EXPECTED_BORDER_SZ:
                            border_issues.append(f"Cell [{ri},{ci}] {side}: sz={sz!r}, expected {EXPECTED_BORDER_SZ!r}")
                        if color is None or color.upper() != EXPECTED_BORDER_COLOR.upper():
                            border_issues.append(f"Cell [{ri},{ci}] {side}: color={color!r}, expected {EXPECTED_BORDER_COLOR!r}")
                        if val != EXPECTED_BORDER_VAL:
                            border_issues.append(f"Cell [{ri},{ci}] {side}: val={val!r}, expected {EXPECTED_BORDER_VAL!r}")

        if not border_issues:
            print(f"PASS: Component 4 — All cell borders are 0.75pt solid dark green (#006400) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 — Border issues ({len(border_issues)} problems):")
            for issue in border_issues[:10]:  # show first 10 issues
                print(f"  {issue}")
            if len(border_issues) > 10:
                print(f"  ... and {len(border_issues) - 10} more issues")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
