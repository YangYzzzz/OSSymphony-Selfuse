"""
Reward Script: Parent Permission Slip for School Field Trip
Task ID: writer_wf_074
Domain: libreoffice_writer
Scoring:
  C1: Title bold+centered "PARENT/GUARDIAN PERMISSION FORM" (0.20)
  C2: School name "Lincoln Middle School" present (0.10)
  C3: Trip details - destination, date, times, transport, cost (0.20)
  C4: Trip purpose section (0.10)
  C5: Student info fields - name, grade, teacher (0.15)
  C6: Medical/emergency fields (0.10)
  C7: Consent paragraph + signature/date line (0.15)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_074'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def get_full_text(doc):
    """Get all text from paragraphs as a single lowercase string."""
    return '\n'.join(p.text for p in doc.paragraphs).lower()


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_text = get_full_text(doc)
    num_paras = len(doc.paragraphs)

    # Gate: document must have meaningful content (at least 5 paragraphs)
    if num_paras < 5:
        print(f"FAIL: Document has only {num_paras} paragraphs — too few for a permission form")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title "PARENT/GUARDIAN PERMISSION FORM" bold + centered (0.20 points)
    try:
        title_found = False
        for para in doc.paragraphs:
            text_lower = para.text.strip().lower()
            if 'parent' in text_lower and ('guardian' in text_lower or 'permission' in text_lower) and 'form' in text_lower:
                # Check bold
                has_bold = any(r.bold for r in para.runs if r.text.strip())
                # Check centered
                is_centered = para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                if has_bold and is_centered:
                    print(f"PASS: Component 1 — Title '{para.text.strip()}' is bold and centered (0.20 pts)")
                    total_score += 0.20
                    title_found = True
                elif has_bold:
                    print(f"PARTIAL: Component 1 — Title is bold but not centered (0.10 pts)")
                    total_score += 0.10
                    title_found = True
                elif is_centered:
                    print(f"PARTIAL: Component 1 — Title is centered but not bold (0.10 pts)")
                    total_score += 0.10
                    title_found = True
                else:
                    print(f"PARTIAL: Component 1 — Title text found but neither bold nor centered (0.05 pts)")
                    total_score += 0.05
                    title_found = True
                break
        if not title_found:
            print("FAIL: Component 1 — Title 'PARENT/GUARDIAN PERMISSION FORM' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: School name "Lincoln Middle School" present (0.10 points)
    try:
        if 'lincoln middle school' in all_text:
            print("PASS: Component 2 — 'Lincoln Middle School' found (0.10 pts)")
            total_score += 0.10
        else:
            print("FAIL: Component 2 — 'Lincoln Middle School' not found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Trip details section (0.20 points)
    # Must include: destination "City Science Museum", date, departure/return times, transportation, cost
    try:
        trip_score = 0.0
        checks = {
            'destination (City Science Museum)': 'city science museum' in all_text,
            'date': bool(re.search(r'date[:\s]', all_text)),
            'departure time': bool(re.search(r'departure\s*(time)?[:\s]', all_text)),
            'return time': bool(re.search(r'return\s*(time)?[:\s]', all_text)),
            'transportation': 'transportation' in all_text or 'transport' in all_text,
            'cost': 'cost' in all_text or '$' in all_text or 'fee' in all_text,
        }
        passed = sum(1 for v in checks.values() if v)
        trip_score = 0.20 * (passed / len(checks))
        for name, result in checks.items():
            status = "ok" if result else "MISSING"
            print(f"  C3 sub-check '{name}': {status}")
        if trip_score > 0:
            print(f"PASS: Component 3 — Trip details ({passed}/{len(checks)} items) ({trip_score:.2f} pts)")
            total_score += trip_score
        else:
            print("FAIL: Component 3 — No trip detail items found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Trip purpose section (0.10 points)
    try:
        # Look for a paragraph describing the purpose of the trip (longer text, mentions museum/science/curriculum)
        purpose_found = False
        for para in doc.paragraphs:
            text_lower = para.text.strip().lower()
            # A purpose paragraph should be substantive (>50 chars) and mention the trip context
            if len(text_lower) > 50 and ('museum' in text_lower or 'science' in text_lower or 'curriculum' in text_lower or 'purpose' in text_lower):
                purpose_found = True
                break
        # Also check for a "Trip Purpose" or "Purpose" heading
        has_purpose_heading = any('purpose' in p.text.lower() for p in doc.paragraphs)

        if purpose_found and has_purpose_heading:
            print("PASS: Component 4 — Trip purpose section with heading found (0.10 pts)")
            total_score += 0.10
        elif purpose_found:
            print("PARTIAL: Component 4 — Trip purpose description found but no heading (0.07 pts)")
            total_score += 0.07
        elif has_purpose_heading:
            print("PARTIAL: Component 4 — Purpose heading found but no descriptive paragraph (0.03 pts)")
            total_score += 0.03
        else:
            print("FAIL: Component 4 — No trip purpose section found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Student info fields - student name, grade, teacher (0.15 points)
    try:
        student_score = 0.0
        has_student_name = bool(re.search(r'student\s*name', all_text))
        has_grade = bool(re.search(r'grade[:\s_]', all_text))
        has_teacher = bool(re.search(r'teacher[:\s_]', all_text))

        fields_found = sum([has_student_name, has_grade, has_teacher])
        student_score = 0.15 * (fields_found / 3)
        print(f"  C5: student_name={'ok' if has_student_name else 'MISSING'}, grade={'ok' if has_grade else 'MISSING'}, teacher={'ok' if has_teacher else 'MISSING'}")
        if student_score > 0:
            print(f"PASS: Component 5 — Student info fields ({fields_found}/3) ({student_score:.2f} pts)")
            total_score += student_score
        else:
            print("FAIL: Component 5 — No student info fields found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Medical/emergency fields (0.10 points)
    try:
        med_score = 0.0
        has_medical = bool(re.search(r'medical|allerg', all_text))
        has_emergency_name = bool(re.search(r'emergency\s*contact\s*(name)?', all_text))
        has_emergency_phone = bool(re.search(r'(emergency.*phone|phone.*emergency|contact\s*phone)', all_text))

        fields_found = sum([has_medical, has_emergency_name, has_emergency_phone])
        med_score = 0.10 * (fields_found / 3)
        print(f"  C6: medical={'ok' if has_medical else 'MISSING'}, emergency_name={'ok' if has_emergency_name else 'MISSING'}, emergency_phone={'ok' if has_emergency_phone else 'MISSING'}")
        if med_score > 0:
            print(f"PASS: Component 6 — Medical/emergency fields ({fields_found}/3) ({med_score:.2f} pts)")
            total_score += med_score
        else:
            print("FAIL: Component 6 — No medical/emergency fields found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Consent paragraph + signature/date line (0.15 points)
    try:
        consent_score = 0.0
        # Consent paragraph: mentions permission/consent/authorize and parent/guardian
        has_consent = bool(re.search(r'(permission|consent|authorize|hereby)', all_text) and
                          re.search(r'(parent|guardian)', all_text))
        # Signature line
        has_signature = bool(re.search(r'signature[:\s_]', all_text))
        # Date line (near signature context)
        has_date_line = bool(re.search(r'(signature.*date|date.*___)', all_text))

        if has_consent and has_signature:
            consent_score = 0.15
            print(f"PASS: Component 7 — Consent paragraph and signature line found (0.15 pts)")
        elif has_consent:
            consent_score = 0.10
            print(f"PARTIAL: Component 7 — Consent paragraph found but no signature line (0.10 pts)")
        elif has_signature:
            consent_score = 0.05
            print(f"PARTIAL: Component 7 — Signature line found but no consent paragraph (0.05 pts)")
        else:
            print("FAIL: Component 7 — No consent paragraph or signature line found")
        total_score += consent_score
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
