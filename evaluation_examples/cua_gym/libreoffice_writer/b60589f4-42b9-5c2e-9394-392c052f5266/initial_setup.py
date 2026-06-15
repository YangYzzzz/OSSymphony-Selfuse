"""
Initial Setup: Page orientation task — document in landscape (needs to be set back to portrait)
Task ID: writer_page_066
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'wide_report'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # ---- Page Setup: A4 Landscape, all margins 2.54 cm ----
    section = doc.sections[0]
    # A4: 210mm x 297mm. In landscape: width=297mm, height=210mm
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ---- Page 1: Cover / Executive Summary ----
    title = doc.add_heading('Quarterly Business Performance Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph('Q1 2025 — Executive Summary')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph('')

    doc.add_heading('Overview', level=1)
    doc.add_paragraph(
        'This report presents a comprehensive analysis of the company\'s operational and financial '
        'performance for the first quarter of 2025. Key metrics across all business units have been '
        'consolidated to provide leadership with a clear picture of progress against annual targets.'
    )

    doc.add_paragraph(
        'Despite global supply chain disruptions, the company achieved a 12.4% revenue growth '
        'year-over-year, exceeding the projected 9% target. Operating margins improved by 1.8 '
        'percentage points to reach 18.3%, driven primarily by efficiency gains in the manufacturing '
        'and logistics divisions.'
    )

    doc.add_heading('Key Highlights', level=2)
    highlights = [
        'Total revenue: $84.7 million (vs. $75.4 million in Q1 2024)',
        'Net profit: $15.5 million, up 21% year-over-year',
        'Customer acquisition: 3,240 new accounts in the quarter',
        'Employee headcount: 1,872 FTE across all regions',
        'Product launches: 4 new SKUs introduced to market',
    ]
    for item in highlights:
        doc.add_paragraph(item, style='List Bullet')

    # Page break to page 2
    doc.add_page_break()

    # ---- Page 2: Financial Performance ----
    doc.add_heading('Financial Performance', level=1)
    doc.add_paragraph(
        'The finance team reports strong performance across all three business segments. '
        'The Technology Solutions division led growth with a 19.2% increase in bookings, '
        'while the Professional Services unit maintained steady margins at 22.1%.'
    )

    # Revenue table
    doc.add_heading('Revenue by Division (USD millions)', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Division', 'Q1 2024', 'Q1 2025', 'YoY Growth']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True

    data = [
        ('Technology Solutions', '35.2', '41.9', '+19.2%'),
        ('Professional Services', '24.8', '27.4', '+10.5%'),
        ('Consumer Products', '15.4', '15.4', '0.0%'),
        ('Total', '75.4', '84.7', '+12.4%'),
    ]
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val

    doc.add_paragraph('')
    doc.add_paragraph(
        'Operating expenses for Q1 2025 totaled $69.2 million, an increase of 10.1% from '
        'the prior year period. R&D expenditure accounted for $8.4 million, representing '
        '9.9% of total revenue, consistent with our investment strategy for product innovation.'
    )

    doc.add_heading('Cost Structure Analysis', level=2)
    doc.add_paragraph(
        'The cost of goods sold (COGS) decreased as a percentage of revenue from 51.3% to '
        '49.7%, reflecting the benefits of renegotiated supplier contracts and improved '
        'production yields. Administrative costs were held flat at $11.2 million through '
        'disciplined budget management across all departments.'
    )

    # Page break to page 3
    doc.add_page_break()

    # ---- Page 3: Outlook & Recommendations ----
    doc.add_heading('Strategic Outlook & Recommendations', level=1)
    doc.add_paragraph(
        'Looking ahead to Q2 2025, the management team anticipates continued revenue momentum '
        'supported by a robust pipeline of $112 million in qualified opportunities. The sales '
        'organization has been expanded with 47 new hires in January and February, primarily '
        'focused on enterprise accounts in the APAC and EMEA regions.'
    )

    doc.add_heading('Investment Priorities', level=2)
    priorities = [
        'Cloud infrastructure expansion — budgeted at $6.2 million for H1 2025',
        'Customer success platform upgrade — $1.8 million, deployment in April',
        'Manufacturing automation (Phase II) — $4.5 million, targeting 8% COGS reduction',
        'Brand refresh and marketing campaign — $2.1 million across digital and event channels',
    ]
    for item in priorities:
        doc.add_paragraph(item, style='List Number')

    doc.add_heading('Risk Factors', level=2)
    doc.add_paragraph(
        'The primary risks identified by the risk management committee include: currency '
        'exchange fluctuations affecting international revenues (estimated exposure of $3.2M), '
        'potential component shortages in the semiconductor supply chain, and regulatory changes '
        'in key markets. Mitigation strategies are documented in the accompanying Risk Register.'
    )

    doc.add_heading('Conclusion', level=2)
    doc.add_paragraph(
        'The company is well-positioned to achieve its full-year 2025 targets of $340 million '
        'in revenue and a 19% operating margin. Continued focus on operational excellence, '
        'customer retention, and disciplined capital allocation will be critical success factors '
        'in the quarters ahead.'
    )

    doc.add_paragraph('')
    closing = doc.add_paragraph('Prepared by: Finance & Strategy Team  |  Approved by: CFO Office')
    closing.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in closing.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
