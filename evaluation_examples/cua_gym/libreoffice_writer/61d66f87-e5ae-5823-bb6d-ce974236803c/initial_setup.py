"""
Initial Setup: Monthly Community Newsletter with mixed formatting in paragraph 2
Task ID: writer_txtfmt_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_041'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Paragraph 0: Title ---
    title = doc.add_heading('Monthly Community Newsletter', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Paragraph 1 (index 1): Subtitle / date ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = subtitle.add_run('April 2025 Edition')
    r.bold = True
    r.font.size = Pt(12)

    # --- Paragraph 2 (index 2): The target paragraph with mixed formatting ---
    # Task instruction: "second paragraph" with mixed bold/italic/underline/color
    para2 = doc.add_paragraph()
    # "Join us for the annual "
    r1 = para2.add_run('Join us for the annual ')

    # "Spring Festival" - bold + red
    r2 = para2.add_run('Spring Festival')
    r2.bold = True
    r2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # " on "
    r3 = para2.add_run(' on ')

    # "April 12th" - italic + underline
    r4 = para2.add_run('April 12th')
    r4.italic = True
    r4.underline = True

    # " at "
    r5 = para2.add_run(' at ')

    # "Central Park" - bold + italic + blue
    r6 = para2.add_run('Central Park')
    r6.bold = True
    r6.italic = True
    r6.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    # ". Activities include "
    r7 = para2.add_run('. Activities include ')

    # "live music" - underline
    r8 = para2.add_run('live music')
    r8.underline = True

    # ", "
    r9 = para2.add_run(', ')

    # "food vendors" - bold
    r10 = para2.add_run('food vendors')
    r10.bold = True

    # ", "
    r11 = para2.add_run(', ')

    # "craft booths" - italic + red
    r12 = para2.add_run('craft booths')
    r12.italic = True
    r12.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # ", and a "
    r13 = para2.add_run(', and a ')

    # "charity raffle" - bold + underline + blue
    r14 = para2.add_run('charity raffle')
    r14.bold = True
    r14.underline = True
    r14.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    # "."
    r15 = para2.add_run('.')

    # --- Paragraph 3: Upcoming Events section header ---
    events_header = doc.add_paragraph()
    r = events_header.add_run('Upcoming Events')
    r.bold = True
    r.font.size = Pt(13)

    # --- Paragraph 4 ---
    doc.add_paragraph(
        'The Community Center will host a series of workshops throughout the month, '
        'covering topics such as gardening, cooking, and digital literacy. '
        'All residents are welcome to attend free of charge.'
    )

    # --- Paragraph 5: Volunteer section ---
    vol_header = doc.add_paragraph()
    r = vol_header.add_run('Volunteer Spotlight')
    r.bold = True
    r.font.size = Pt(13)

    # --- Paragraph 6 ---
    doc.add_paragraph(
        'This month we recognize Sarah Mitchell and David Torres for their outstanding '
        'contributions to the neighborhood clean-up initiative on March 22nd. '
        'Their dedication has made our community a cleaner and more welcoming place.'
    )

    # --- Paragraph 7: Contact info ---
    contact = doc.add_paragraph()
    r = contact.add_run('Contact us at: ')
    r.bold = True
    contact.add_run('newsletter@community.org | (555) 867-5309')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
