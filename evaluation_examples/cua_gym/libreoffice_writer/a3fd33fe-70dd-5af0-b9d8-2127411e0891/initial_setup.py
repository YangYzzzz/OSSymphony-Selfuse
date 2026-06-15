"""
Initial Setup: Insert company logo into letterhead document
Task ID: writer_obj_056
Domain: libreoffice_writer

Creates:
  - /home/user/Desktop/letterhead.docx  -- formal letter template with NO logo
  - /home/user/Desktop/logo.png         -- 500x500 pixel company logo image
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_056'
LETTERHEAD_PATH = f'{WORKDIR}/letterhead.docx'
LOGO_PATH = f'{WORKDIR}/logo.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_logo():
    """Create a 500x500 pixel company logo PNG."""
    img = Image.new('RGB', (500, 500), color=(41, 128, 185))  # corporate blue

    draw = ImageDraw.Draw(img)

    # Draw a simple geometric logo: white rectangle with "TC" text
    draw.rectangle([50, 50, 450, 450], fill=(255, 255, 255), outline=(41, 128, 185), width=0)
    draw.rectangle([80, 80, 420, 420], fill=(41, 128, 185))
    draw.rectangle([110, 110, 390, 390], fill=(255, 255, 255))

    # Draw "TC" letters for TechCorp
    draw.rectangle([130, 150, 240, 180], fill=(41, 128, 185))   # T top bar
    draw.rectangle([175, 180, 200, 340], fill=(41, 128, 185))   # T stem
    draw.rectangle([260, 150, 370, 180], fill=(41, 128, 185))   # C top
    draw.rectangle([260, 310, 370, 340], fill=(41, 128, 185))   # C bottom
    draw.rectangle([260, 180, 290, 310], fill=(41, 128, 185))   # C left

    # Add a thin border ring
    draw.ellipse([60, 60, 440, 440], outline=(41, 128, 185), width=5)

    img.save(LOGO_PATH, 'PNG')
    print(f'Logo created: {LOGO_PATH}')


def create_letterhead():
    """Create a formal letter template WITHOUT any logo."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.0)

    # Company name heading (top of letterhead)
    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = heading.add_run('TechCorp Solutions, Inc.')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x96)  # corporate blue

    # Company address block
    address_para = doc.add_paragraph()
    address_run = address_para.add_run(
        '1200 Innovation Drive, Suite 400\n'
        'San Francisco, CA 94105\n'
        'Tel: +1 (415) 555-0198  |  Fax: +1 (415) 555-0199\n'
        'www.techcorpsolutions.com  |  contact@techcorp.com'
    )
    address_run.font.size = Pt(10)
    address_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Horizontal rule (simulated with underscored paragraph)
    rule_para = doc.add_paragraph('_' * 72)
    rule_para.paragraph_format.space_before = Pt(0)
    rule_para.paragraph_format.space_after = Pt(6)

    # Date line
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(12)
    date_para.paragraph_format.space_after = Pt(12)
    date_run = date_para.add_run('March 5, 2025')
    date_run.font.size = Pt(11)

    # Recipient block
    recipient_block = [
        'Dr. Amanda Richardson',
        'Chief Technology Officer',
        'Nexus Innovations Ltd.',
        '350 Market Street, Floor 12',
        'San Francisco, CA 94105',
    ]
    for line in recipient_block:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(0)

    # Salutation
    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_before = Pt(18)
    salutation.paragraph_format.space_after = Pt(12)
    sal_run = salutation.add_run('Dear Dr. Richardson,')
    sal_run.font.size = Pt(11)

    # Body paragraphs
    body_text = [
        (
            'I am writing on behalf of TechCorp Solutions, Inc. to formally propose a strategic '
            'partnership in the development of next-generation cloud infrastructure services. '
            'We have closely followed Nexus Innovations\' recent work in distributed systems '
            'and believe there is substantial synergy between our organizations.'
        ),
        (
            'Our team has developed a proprietary edge-computing platform, TechEdge Pro, which '
            'has demonstrated a 40% reduction in latency for enterprise workloads in pilot '
            'deployments with three Fortune 500 clients in Q4 2024. We are confident this '
            'technology, combined with Nexus\'s expertise in AI-driven optimization, could '
            'create a compelling joint offering for the financial services and healthcare sectors.'
        ),
        (
            'We propose an initial meeting to explore the commercial and technical feasibility '
            'of such a collaboration. Our business development team is available at your '
            'convenience during the week of March 17\u201321, 2025. Please let us know a '
            'time that works best for you, and we will arrange all logistics accordingly.'
        ),
        (
            'Thank you for your consideration. We look forward to the possibility of working '
            'together and are excited about the value we can jointly deliver to our customers.'
        ),
    ]

    for body in body_text:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        r = p.add_run(body)
        r.font.size = Pt(11)

    # Closing
    closing_para = doc.add_paragraph()
    closing_para.paragraph_format.space_before = Pt(6)
    closing_para.paragraph_format.space_after = Pt(0)
    closing_run = closing_para.add_run('Sincerely yours,')
    closing_run.font.size = Pt(11)

    # Signature lines
    sig_lines = [
        '',
        '',
        'Jonathan K. Whitfield',
        'Chief Executive Officer',
        'TechCorp Solutions, Inc.',
        'Direct: +1 (415) 555-0201',
        'jwhitfield@techcorp.com',
    ]
    for line in sig_lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(0)

    doc.save(LETTERHEAD_PATH)
    print(f'Letterhead created: {LETTERHEAD_PATH}')


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)
    create_logo()
    create_letterhead()

    # GUI-ready startup: open letterhead in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{LETTERHEAD_PATH}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
