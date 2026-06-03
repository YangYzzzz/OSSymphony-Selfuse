"""
Reward Script: Autoformat customer reference table in LibreOffice Writer
Task ID: writer_mktg_046
Domain: libreoffice_writer
Scoring:
  Component 1: Header row bold text + dark blue background (#1B3A5C) + white text  (0.30 pts)
  Component 2: Alternating row shading (white / light gray #F5F5F5)                (0.25 pts)
  Component 3: Annual Revenue column (col 3) right-aligned in data rows             (0.20 pts)
  Component 4: NPS Score column (col 4) center-aligned in data rows                 (0.10 pts)
  Component 5: Summary row with 'Total: 15 customers' (bold) and total revenue
               '$4,405,450' (bold, right-aligned)                                   (0.15 pts)
  Total: 1.0
"""

import os
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_mktg_046'

def get_cell_fill(cell):
    """Return the fill hex string of a cell, or None if not set."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    return fill  # e.g. '1B3A5C', 'FFFFFF', 'F5F5F5', or 'auto'

def get_cell_runs_bold_and_color(cell):
    """
    Return (has_bold, has_white_text) for non-empty runs in a cell.
    has_bold: at least one non-empty run with bold=True
    has_white_text: at least one non-empty run with color close to white (FFFFFF)
    """
    has_bold = False
    has_white = False
    for para in cell.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            if run.font.bold is True:
                has_bold = True
            try:
                rgb = run.font.color.rgb
                if rgb is not None:
                    # RGBColor supports indexing: rgb[0], rgb[1], rgb[2] -> 0-255 integers
                    r, g, b = rgb[0], rgb[1], rgb[2]
                    if r >= 200 and g >= 200 and b >= 200:
                        has_white = True
            except Exception:
                pass
    return has_bold, has_white

def color_distance(hex1, hex2):
    """Simple Euclidean distance between two hex RGB colors."""
    try:
        r1, g1, b1 = int(hex1[0:2], 16), int(hex1[2:4], 16), int(hex1[4:6], 16)
        r2, g2, b2 = int(hex2[0:2], 16), int(hex2[2:4], 16), int(hex2[4:6], 16)
        return ((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2) ** 0.5
    except Exception:
        return 999.0

def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    print(f"INFO: Table has {num_rows} rows and {num_cols} columns")

    # ------------------------------------------------------------------
    # Component 1: Header row — bold text + dark blue background + white text (0.30 pts)
    # Expected: fill=1B3A5C, bold text, white color in all 5 header cells
    # FAILS on initial (no shading, no bold), PASSES on golden
    # ------------------------------------------------------------------
    try:
        header_row = table.rows[0]
        header_bg_ok = 0
        header_bold_ok = 0
        header_white_ok = 0
        target_blue = "1B3A5C"

        for j, cell in enumerate(header_row.cells):
            fill = get_cell_fill(cell)
            if fill and color_distance(fill.upper(), target_blue.upper()) < 30:
                header_bg_ok += 1
            has_bold, has_white = get_cell_runs_bold_and_color(cell)
            if has_bold:
                header_bold_ok += 1
            if has_white:
                header_white_ok += 1

        # Require at least 4/5 cells to pass each check for partial tolerance
        bg_pass = header_bg_ok >= 4
        bold_pass = header_bold_ok >= 4
        white_pass = header_white_ok >= 4

        if bg_pass and bold_pass and white_pass:
            print(f"PASS: Component 1 — Header row: dark blue bg ({header_bg_ok}/5), bold ({header_bold_ok}/5), white text ({header_white_ok}/5) (0.30 pts)")
            total_score += 0.30
        elif bg_pass or bold_pass:
            # Partial: background set but text not properly formatted or vice versa
            print(f"PARTIAL: Component 1 — Header bg_ok={header_bg_ok}/5, bold_ok={header_bold_ok}/5, white_ok={header_white_ok}/5 (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Header row not formatted: bg_ok={header_bg_ok}/5, bold_ok={header_bold_ok}/5, white_ok={header_white_ok}/5")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ------------------------------------------------------------------
    # Component 2: Alternating row shading — odd data rows white, even data rows light gray (0.25 pts)
    # Data rows are rows 1-15 (indices 1 to 15 in the table)
    # odd indices (1, 3, 5, ...) -> white (#FFFFFF)
    # even indices (2, 4, 6, ...) -> light gray (#F5F5F5)
    # FAILS on initial (no shading), PASSES on golden
    # ------------------------------------------------------------------
    try:
        if num_rows < 16:
            print(f"FAIL: Component 2 — Not enough data rows for alternating shading ({num_rows} rows)")
        else:
            white_target = "FFFFFF"
            gray_target = "F5F5F5"
            correct_shading = 0
            total_data_rows = 15

            for r_idx in range(1, 16):
                row = table.rows[r_idx]
                cell = row.cells[0]  # Check first cell of each data row
                fill = get_cell_fill(cell)
                if fill is None:
                    fill = "NONE"

                # r_idx 1,3,5,... (odd) -> white; r_idx 2,4,6,... (even) -> gray
                if r_idx % 2 == 1:
                    expected = white_target
                else:
                    expected = gray_target

                if fill and color_distance(fill.upper(), expected.upper()) < 20:
                    correct_shading += 1

            ratio = correct_shading / total_data_rows
            if ratio >= 0.9:
                print(f"PASS: Component 2 — Alternating shading: {correct_shading}/{total_data_rows} rows correct (0.25 pts)")
                total_score += 0.25
            elif ratio >= 0.5:
                print(f"PARTIAL: Component 2 — Alternating shading: {correct_shading}/{total_data_rows} rows correct (0.12 pts)")
                total_score += 0.12
            else:
                print(f"FAIL: Component 2 — Alternating shading: only {correct_shading}/{total_data_rows} rows correct")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ------------------------------------------------------------------
    # Component 3: Annual Revenue column (col index 3) right-aligned in data rows (0.20 pts)
    # Checks rows 1-15; header alignment not required to be right
    # FAILS on initial (align=None), PASSES on golden (align=RIGHT)
    # ------------------------------------------------------------------
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        if num_rows < 2 or num_cols < 4:
            print("FAIL: Component 3 — Not enough rows/cols for revenue alignment check")
        else:
            right_aligned = 0
            check_rows = min(15, num_rows - 1)
            for r_idx in range(1, check_rows + 1):
                row = table.rows[r_idx]
                cell = row.cells[3]
                for para in cell.paragraphs:
                    if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                        right_aligned += 1
                        break

            ratio = right_aligned / check_rows
            if ratio >= 0.9:
                print(f"PASS: Component 3 — Revenue column right-aligned: {right_aligned}/{check_rows} rows (0.20 pts)")
                total_score += 0.20
            elif ratio >= 0.5:
                print(f"PARTIAL: Component 3 — Revenue column right-aligned: {right_aligned}/{check_rows} rows (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 3 — Revenue column not right-aligned: only {right_aligned}/{check_rows} rows")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ------------------------------------------------------------------
    # Component 4: NPS Score column (col index 4) center-aligned in data rows (0.10 pts)
    # FAILS on initial (align=None), PASSES on golden (align=CENTER)
    # ------------------------------------------------------------------
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        if num_rows < 2 or num_cols < 5:
            print("FAIL: Component 4 — Not enough rows/cols for NPS alignment check")
        else:
            center_aligned = 0
            check_rows = min(15, num_rows - 1)
            for r_idx in range(1, check_rows + 1):
                row = table.rows[r_idx]
                cell = row.cells[4]
                for para in cell.paragraphs:
                    if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                        center_aligned += 1
                        break

            ratio = center_aligned / check_rows
            if ratio >= 0.9:
                print(f"PASS: Component 4 — NPS column centered: {center_aligned}/{check_rows} rows (0.10 pts)")
                total_score += 0.10
            elif ratio >= 0.5:
                print(f"PARTIAL: Component 4 — NPS column centered: {center_aligned}/{check_rows} rows (0.05 pts)")
                total_score += 0.05
            else:
                print(f"FAIL: Component 4 — NPS column not centered: only {center_aligned}/{check_rows} rows")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ------------------------------------------------------------------
    # Component 5: Summary row — 'Total: 15 customers' (bold) in col 0 and
    # total revenue '$4,405,450' (bold, right-aligned) in col 3 (0.15 pts)
    # FAILS on initial (only 16 rows, no summary row), PASSES on golden
    # ------------------------------------------------------------------
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        if num_rows < 17:
            print(f"FAIL: Component 5 — No summary row found (table has {num_rows} rows, expected 17)")
        else:
            summary_row = table.rows[16]
            col0_cell = summary_row.cells[0]
            col3_cell = summary_row.cells[3]

            # Check col 0: contains customer count text and is bold
            col0_text = col0_cell.text.strip()
            col0_has_count = "15" in col0_text and "customer" in col0_text.lower()
            col0_bold = any(
                run.font.bold is True
                for para in col0_cell.paragraphs
                for run in para.runs
                if run.text.strip()
            )

            # Check col 3: contains total revenue value and is bold + right-aligned
            col3_text = col3_cell.text.strip()
            # Accept $4,405,450 or similar (sum of the 15 revenues)
            expected_total = "4,405,450"
            col3_has_total = expected_total in col3_text.replace(",", ",")
            col3_bold = any(
                run.font.bold is True
                for para in col3_cell.paragraphs
                for run in para.runs
                if run.text.strip()
            )
            col3_right_aligned = any(
                para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT
                for para in col3_cell.paragraphs
            )

            count_ok = col0_has_count and col0_bold
            total_ok = col3_has_total and col3_bold

            if count_ok and total_ok:
                print(f"PASS: Component 5 — Summary row: count='{col0_text}' (bold={col0_bold}), revenue='{col3_text}' (bold={col3_bold}) (0.15 pts)")
                total_score += 0.15
            elif count_ok or total_ok:
                print(f"PARTIAL: Component 5 — Summary row partial: count_ok={count_ok}, total_ok={total_ok} (0.08 pts)")
                print(f"  col0: text={repr(col0_text)}, has_count={col0_has_count}, bold={col0_bold}")
                print(f"  col3: text={repr(col3_text)}, has_total={col3_has_total}, bold={col3_bold}")
                total_score += 0.08
            else:
                print(f"FAIL: Component 5 — Summary row missing or incorrect")
                print(f"  col0: text={repr(col0_text)}, has_count={col0_has_count}, bold={col0_bold}")
                print(f"  col3: text={repr(col3_text)}, has_total={col3_has_total}, bold={col3_bold}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM env
file_path = f'{WORKDIR}/customer_reference_list.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
