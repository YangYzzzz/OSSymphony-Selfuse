"""
Initial Setup: Customer Reference List - Plain Table (No Formatting)
Task ID: writer_mktg_046
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_046'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/customer_reference_list.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Title paragraph
    title = doc.add_paragraph()
    title_run = title.add_run('Customer Reference List \u2014 Q1 2026')
    title_run.font.size = Pt(16)
    title_run.bold = True
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(12)

    # Data rows (15 customers) — no formatting, plain table
    customers = [
        ('Apex Solutions LLC',        'Technology',       '2024-01-15', '$245,000', '72'),
        ('Hartwell & Partners',       'Legal Services',   '2023-09-01', '$188,500', '65'),
        ('Bright Horizon Media',      'Marketing',        '2024-03-20', '$97,200',  '81'),
        ('Cascade Industrial',        'Manufacturing',    '2022-11-08', '$512,000', '58'),
        ('Delphi Financial Group',    'Finance',          '2023-06-14', '$334,750', '74'),
        ('EcoVerde Consulting',       'Environmental',    '2024-02-28', '$76,400',  '89'),
        ('Fortuna Retail Holdings',   'Retail',           '2022-07-03', '$421,300', '63'),
        ('Granite Peak Engineering',  'Engineering',      '2023-04-17', '$298,000', '70'),
        ('Harbor View Hospitality',   'Hospitality',      '2023-12-01', '$163,800', '77'),
        ('Innovatech Systems',        'Technology',       '2024-05-10', '$289,600', '85'),
        ('Juniper Health Network',    'Healthcare',       '2022-08-22', '$457,200', '69'),
        ('Kinsley Logistics Corp',    'Logistics',        '2023-10-05', '$215,400', '61'),
        ('Luminary Education Group',  'Education',        '2024-04-12', '$88,700',  '83'),
        ('Meridian Asset Management', 'Finance',          '2022-12-19', '$623,100', '55'),
        ('NovaBridge Pharma',         'Pharmaceuticals',  '2023-07-30', '$394,500', '78'),
    ]

    headers = ['Customer Name', 'Industry', 'Contract Start', 'Annual Revenue', 'NPS Score']

    # Create table: 1 header + 15 data rows = 16 rows, 5 columns
    table = doc.add_table(rows=16, cols=5)
    table.style = 'Table Grid'

    # Set equal column widths (approximately)
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(1.5)

    # Header row — plain, no formatting
    header_row = table.rows[0]
    for j, h in enumerate(headers):
        cell = header_row.cells[j]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(h)
        # No bold, no color — plain header as imported from CRM
        run.font.size = Pt(11)

    # Data rows — plain text, left-aligned
    for i, (name, industry, contract_start, revenue, nps) in enumerate(customers):
        row = table.rows[i + 1]
        row_data = [name, industry, contract_start, revenue, nps]
        for j, value in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(value)
            run.font.size = Pt(11)
            # All cells left-aligned (default) — no special alignment

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
