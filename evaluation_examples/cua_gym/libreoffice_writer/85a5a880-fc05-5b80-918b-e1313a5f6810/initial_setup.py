"""
Initial Setup: Office memo template with page border on all four sides.
Task ID: writer_page_038
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'memo_template'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_borders(section, sides):
    """Add page borders to specified sides of a section.

    Args:
        section: docx Section object
        sides: list of side names, e.g. ['top', 'left', 'bottom', 'right']
    """
    sectPr = section._sectPr

    # Remove existing pgBorders if present
    existing = sectPr.find(qn('w:pgBorders'))
    if existing is not None:
        sectPr.remove(existing)

    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')

    # Borders must appear in correct order per OOXML spec
    for side in ['top', 'left', 'bottom', 'right']:
        if side in sides:
            border_el = OxmlElement(f'w:{side}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '8')       # 1pt = 8 eighths-of-a-point
            border_el.set(qn('w:space'), '24')   # standard padding
            border_el.set(qn('w:color'), '000000')
            pgBorders.append(border_el)

    # Insert pgBorders before pgMar (after pgSz) per OOXML schema order
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is not None:
        pgSz.addnext(pgBorders)
    else:
        sectPr.insert(0, pgBorders)


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()
    section = doc.sections[0]

    # A4 portrait, 2.54cm margins on all sides
    section.page_width = int(Cm(21))
    section.page_height = int(Cm(29.7))
    section.left_margin = int(Cm(2.54))
    section.right_margin = int(Cm(2.54))
    section.top_margin = int(Cm(2.54))
    section.bottom_margin = int(Cm(2.54))

    # Page border on ALL FOUR sides (initial state — task requires removing top/bottom)
    add_page_borders(section, ['top', 'left', 'bottom', 'right'])

    # --- Company header ---
    heading_para = doc.add_paragraph()
    heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading_para.add_run('MERIDIAN CONSULTING GROUP')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x3D, 0x7A)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = sub_para.add_run('Internal Memorandum')
    sub_run.font.size = Pt(12)
    sub_run.italic = True

    doc.add_paragraph()  # spacer

    # --- Memo header fields ---
    def add_memo_field(label, value):
        para = doc.add_paragraph()
        label_run = para.add_run(f'{label}:  ')
        label_run.bold = True
        label_run.font.size = Pt(11)
        value_run = para.add_run(value)
        value_run.font.size = Pt(11)
        return para

    add_memo_field('DATE', 'March 5, 2025')
    add_memo_field('TO', 'All Department Heads')
    add_memo_field('FROM', 'Elena Whitmore, Chief Operations Officer')
    add_memo_field('RE', 'Q1 2025 Operational Review and Budget Adjustments')

    # --- Divider ---
    divider = doc.add_paragraph('_' * 65)
    divider.paragraph_format.space_after = Pt(6)

    # --- Body ---
    body_intro = doc.add_paragraph()
    body_intro.paragraph_format.space_before = Pt(6)
    body_intro.add_run(
        'This memorandum summarizes the findings from the Q1 2025 operational review '
        'conducted between February 10 and February 28, 2025. All department heads are '
        'requested to review the information below and respond with corrective action '
        'plans by March 20, 2025.'
    ).font.size = Pt(11)

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2_run = p2.add_run('1.  Financial Performance Overview')
    p2_run.bold = True
    p2_run.font.size = Pt(11)

    p3 = doc.add_paragraph()
    p3.add_run(
        'Total Q1 revenue reached $4.82M against a target of $5.10M, representing a '
        'shortfall of 5.5%. Operating expenses were $3.41M, resulting in an operating '
        'margin of 29.3%. The primary variance drivers were increased contractor costs '
        'in the Technology division ($187,000 over budget) and lower-than-expected '
        'client onboarding fees in the Advisory division.'
    ).font.size = Pt(11)

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4_run = p4.add_run('2.  Staffing and Resource Allocation')
    p4_run.bold = True
    p4_run.font.size = Pt(11)

    p5 = doc.add_paragraph()
    p5.add_run(
        'Headcount at end of Q1 was 143 FTE (full-time equivalents), with 7 open '
        'positions pending approval. The HR department has flagged a 12% increase in '
        'voluntary attrition in the Customer Success team. Recommended actions include '
        'a compensation benchmarking review and updated onboarding procedures for '
        'mid-level roles.'
    ).font.size = Pt(11)

    doc.add_paragraph()

    p6 = doc.add_paragraph()
    p6_run = p6.add_run('3.  Strategic Priorities for Q2 2025')
    p6_run.bold = True
    p6_run.font.size = Pt(11)

    p7 = doc.add_paragraph()
    p7.add_run(
        'The executive team has identified three focus areas for Q2: (a) acceleration '
        'of the digital transformation initiative in Operations, (b) renegotiation of '
        'the Apex Logistics contract to recover margin, and (c) launch of the revised '
        'client engagement framework developed by the Strategy team.'
    ).font.size = Pt(11)

    doc.add_paragraph()

    # --- Action required ---
    action_para = doc.add_paragraph()
    action_run = action_para.add_run('Action Required:  ')
    action_run.bold = True
    action_run.font.size = Pt(11)
    action_para.add_run(
        'Please submit your department\'s response and Q2 action plan to '
        'ewhitmore@meridiancg.com no later than March 20, 2025.'
    ).font.size = Pt(11)

    # --- Footer note ---
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer_para.add_run('CONFIDENTIAL — FOR INTERNAL USE ONLY')
    footer_run.bold = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
