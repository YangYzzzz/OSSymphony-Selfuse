"""
Initial Setup: Thesis document with uniform headers on all pages
Task ID: writer_af_013
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_af_013'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_field(paragraph):
    """Add a PAGE field code to a paragraph."""
    run = paragraph.add_run()
    fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._element.append(fldChar1)
    run2 = paragraph.add_run()
    instrText = run2._element.makeelement(qn('w:instrText'), {})
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)
    run3 = paragraph.add_run()
    fldChar2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run3._element.append(fldChar2)


# Chapter content data - realistic thesis content
CHAPTERS = [
    {
        "title": "Introduction",
        "sections": [
            ("Background and Motivation",
             "The rapid advancement of artificial intelligence and machine learning has fundamentally "
             "transformed numerous industries over the past decade. From healthcare diagnostics to autonomous "
             "vehicles, the integration of neural network architectures into real-world applications has "
             "demonstrated unprecedented capabilities. This thesis examines the theoretical foundations and "
             "practical implications of transformer-based architectures in natural language processing, with "
             "particular emphasis on their application to domain-specific knowledge extraction."),
            ("Research Questions",
             "This research addresses three primary questions. First, how can pre-trained language models be "
             "effectively adapted for specialized scientific domains without catastrophic forgetting of general "
             "knowledge? Second, what architectural modifications improve the model's ability to capture "
             "long-range dependencies in technical documentation? Third, how do different fine-tuning strategies "
             "affect downstream task performance across varying dataset sizes?"),
            ("Thesis Structure",
             "The remainder of this thesis is organized as follows. Chapter 2 provides a comprehensive review "
             "of related work in transfer learning and domain adaptation. Chapter 3 describes our proposed "
             "methodology, including the novel attention mechanism and training procedure. Chapter 4 presents "
             "the experimental setup and results, while Chapter 5 concludes with a discussion of findings "
             "and future research directions."),
            ("Scope and Limitations",
             "This study focuses specifically on English-language scientific publications in the biomedical "
             "domain. While the proposed methods are theoretically applicable to other languages and domains, "
             "empirical validation is limited to the datasets described in Chapter 4. The computational "
             "requirements for training were constrained to a cluster of 8 NVIDIA A100 GPUs, which may "
             "limit the scalability analysis presented."),
        ]
    },
    {
        "title": "Literature Review",
        "sections": [
            ("Transfer Learning Foundations",
             "Transfer learning has emerged as a cornerstone technique in modern machine learning, enabling "
             "models trained on large-scale datasets to be repurposed for specialized tasks with limited "
             "labeled data. The seminal work by Pan and Yang (2010) established a formal taxonomy of transfer "
             "learning approaches, categorizing them into inductive, transductive, and unsupervised settings. "
             "Subsequent developments in deep transfer learning, particularly through pre-trained convolutional "
             "neural networks for computer vision tasks, demonstrated the remarkable generalizability of "
             "learned representations across related domains."),
            ("Transformer Architecture Evolution",
             "The transformer architecture, introduced by Vaswani et al. (2017) in their landmark paper "
             "'Attention Is All You Need,' represented a paradigm shift from recurrent neural networks. "
             "By replacing sequential processing with self-attention mechanisms, transformers achieved superior "
             "parallelization and captured long-range dependencies more effectively. The architecture comprises "
             "an encoder-decoder structure with multi-head attention, position-wise feed-forward networks, "
             "and residual connections with layer normalization."),
            ("Pre-trained Language Models",
             "The development of BERT (Devlin et al., 2019) marked a significant milestone in natural language "
             "processing by demonstrating the effectiveness of bidirectional pre-training. GPT (Radford et al., "
             "2018) and its successors established that autoregressive language modeling at scale could achieve "
             "remarkable few-shot learning capabilities. T5 (Raffel et al., 2020) unified multiple NLP tasks "
             "into a text-to-text framework, while models like RoBERTa (Liu et al., 2019) and ALBERT (Lan et "
             "al., 2020) explored optimization strategies for pre-training efficiency."),
            ("Domain Adaptation in NLP",
             "Domain adaptation for NLP models has been extensively studied, with approaches ranging from "
             "continued pre-training on domain-specific corpora to more sophisticated techniques involving "
             "adversarial training and curriculum learning. BioBERT (Lee et al., 2020) demonstrated significant "
             "improvements on biomedical text mining tasks through domain-specific pre-training on PubMed "
             "abstracts and PMC full-text articles. SciBERT (Beltagy et al., 2019) extended this approach "
             "to the broader scientific literature."),
            ("Knowledge Distillation",
             "Knowledge distillation techniques, pioneered by Hinton et al. (2015), provide a mechanism for "
             "compressing large models into smaller, more efficient architectures while preserving performance. "
             "DistilBERT (Sanh et al., 2019) achieved 97% of BERT's performance with 40% fewer parameters "
             "through careful application of distillation during pre-training. This approach has been extended "
             "to domain-specific settings, where teacher models trained on general corpora guide student "
             "models fine-tuned for specialized tasks."),
        ]
    },
    {
        "title": "Methodology",
        "sections": [
            ("Proposed Architecture",
             "Our proposed architecture extends the standard transformer encoder with a domain-adaptive "
             "attention layer that dynamically adjusts attention patterns based on input characteristics. "
             "The model consists of 12 transformer layers with a hidden dimension of 768 and 12 attention "
             "heads, consistent with the BERT-base configuration. The novel component is a gating mechanism "
             "inserted between the multi-head attention and feed-forward sublayers that learns to weight "
             "domain-general versus domain-specific representations."),
            ("Training Procedure",
             "The training procedure follows a three-phase approach. In Phase 1, the model undergoes masked "
             "language model pre-training on a general English corpus comprising 16GB of text from Wikipedia "
             "and BookCorpus. Phase 2 introduces domain-adaptive pre-training on 4.5 million biomedical "
             "abstracts from PubMed, using a gradually increasing domain mixing ratio. Phase 3 applies "
             "task-specific fine-tuning with a learning rate of 2e-5 and batch size of 32 for a maximum "
             "of 10 epochs with early stopping."),
            ("Attention Mechanism Modifications",
             "The domain-adaptive attention mechanism introduces a learnable gate vector g of dimension d_model "
             "that modulates the attention output before it is added to the residual connection. Formally, "
             "given the standard multi-head attention output A, the gated output is computed as: "
             "A_gated = sigma(g) * A_domain + (1 - sigma(g)) * A_general, where sigma denotes the sigmoid "
             "function. This allows the model to smoothly interpolate between general and domain-specific "
             "attention patterns on a per-layer, per-position basis."),
            ("Dataset Construction",
             "We construct three evaluation datasets from publicly available biomedical resources. The Named "
             "Entity Recognition dataset comprises 18,450 annotated sentences from NCBI Disease, BC5CDR, "
             "and JNLPBA corpora. The Relation Extraction dataset contains 12,800 sentence pairs from "
             "ChemProt and DDI Extraction 2013 shared task. The Question Answering dataset includes 7,200 "
             "questions derived from BioASQ challenges spanning multiple years of competition."),
            ("Evaluation Metrics",
             "Performance is measured using standard metrics appropriate for each task type. For Named Entity "
             "Recognition, we report entity-level precision, recall, and F1 score using exact span matching. "
             "Relation Extraction performance is evaluated using micro-averaged F1 score across all relation "
             "types. Question Answering accuracy is measured using both exact match and token-level F1 scores, "
             "following the conventions established by the SQuAD evaluation protocol."),
        ]
    },
    {
        "title": "Results and Discussion",
        "sections": [
            ("Named Entity Recognition Results",
             "Table 4.1 presents the NER results across all three benchmark datasets. Our proposed model "
             "achieves an average F1 score of 89.4%, representing a 2.1 percentage point improvement over "
             "the BioBERT baseline. The improvement is most pronounced on the JNLPBA dataset (+3.2% F1), "
             "which contains the most diverse entity types including DNA, RNA, cell lines, cell types, "
             "and proteins. Analysis of error patterns reveals that the gated attention mechanism is "
             "particularly effective at disambiguating overlapping entity boundaries."),
            ("Relation Extraction Results",
             "For the Relation Extraction task, our model achieves micro-F1 scores of 78.6% on ChemProt "
             "and 84.2% on DDI 2013, compared to 76.1% and 82.5% respectively for the BioBERT baseline. "
             "The per-class analysis reveals that improvements are concentrated in relation types that "
             "require understanding of long-range contextual dependencies, such as 'substrate' and "
             "'inhibitor' relations in ChemProt that often span multiple clauses within a sentence."),
            ("Question Answering Performance",
             "On the BioASQ benchmark, our model achieves exact match accuracy of 42.8% and F1 score of "
             "61.3%, compared to 39.5% and 58.7% for BioBERT. Qualitative analysis of correctly answered "
             "questions indicates that the domain-adaptive attention mechanism improves the model's ability "
             "to locate relevant evidence spans in longer passages, particularly when the answer requires "
             "synthesizing information from multiple sentences within the context paragraph."),
            ("Ablation Study",
             "To understand the contribution of each component, we conduct a systematic ablation study. "
             "Removing the gating mechanism reduces average F1 by 1.4%, confirming its importance. Using "
             "a fixed mixing ratio instead of learned gates results in a 0.8% decrease. Skipping Phase 2 "
             "(domain-adaptive pre-training) leads to the largest performance drop of 2.6%, underscoring "
             "the value of intermediate domain exposure before task-specific fine-tuning."),
            ("Computational Analysis",
             "The total training time for our three-phase procedure is approximately 72 hours on 8 A100 "
             "GPUs. Phase 1 requires 48 hours, Phase 2 takes 18 hours, and Phase 3 fine-tuning completes "
             "in under 6 hours across all tasks. The additional parameters introduced by the gating "
             "mechanism account for less than 0.5% increase in model size, making the overhead negligible "
             "for both training and inference. Memory consumption during training peaks at 28GB per GPU."),
        ]
    },
    {
        "title": "Conclusion and Future Work",
        "sections": [
            ("Summary of Contributions",
             "This thesis has presented a novel domain-adaptive transformer architecture for biomedical "
             "natural language processing. The key contribution is a gated attention mechanism that enables "
             "smooth interpolation between general and domain-specific representations, achieving state-of-the-art "
             "results across three benchmark tasks. The three-phase training procedure provides a principled "
             "approach to gradually introducing domain knowledge without sacrificing general language "
             "understanding capabilities."),
            ("Implications for Practice",
             "The practical implications of this work are significant for biomedical text mining applications. "
             "The improved NER and RE performance directly benefits automated literature screening, drug "
             "interaction detection, and clinical trial matching systems. The relatively modest computational "
             "overhead of the proposed modifications makes them accessible to research groups without "
             "access to massive computing infrastructure, democratizing advances in biomedical NLP."),
            ("Limitations and Future Directions",
             "Several limitations of this work suggest promising directions for future research. The evaluation "
             "is currently limited to English-language biomedical text; extending to multilingual and cross-lingual "
             "settings would significantly broaden applicability. The gating mechanism could be enhanced with "
             "more sophisticated conditioning, perhaps incorporating explicit domain indicators or metadata. "
             "Additionally, exploring the application of this approach to generative tasks such as biomedical "
             "text summarization and report generation represents an exciting frontier."),
            ("Broader Impact",
             "As AI systems become increasingly integrated into biomedical research workflows, ensuring their "
             "reliability and transparency is paramount. Our work contributes to this goal by providing "
             "interpretable attention patterns through the gating mechanism, allowing researchers to understand "
             "when the model relies on general versus domain-specific knowledge. This interpretability feature "
             "could be valuable for building trust in AI-assisted clinical decision support systems and "
             "accelerating the pace of biomedical discovery."),
        ]
    },
]


def create_initial():
    doc = Document()

    # ---- Page setup ----
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ---- Uniform header with page number on ALL pages ----
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_title = hp.add_run("Domain-Adaptive Transformers for Biomedical NLP  —  Page ")
    run_title.font.size = Pt(9)
    run_title.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    # Add PAGE field
    add_page_number_field(hp)

    # ---- Title page content ----
    title_para = doc.add_heading("Domain-Adaptive Transformer Architectures for Biomedical Natural Language Processing", level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_s = subtitle.add_run("A Thesis Submitted in Partial Fulfillment of the Requirements\n"
                              "for the Degree of Doctor of Philosophy\n\n"
                              "Department of Computer Science\n"
                              "Stanford University\n\n"
                              "Elena Vasquez\n"
                              "March 2026")
    run_s.font.size = Pt(12)

    # Page break after title page
    doc.add_page_break()

    # ---- Abstract ----
    abstract_heading = doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This thesis investigates domain-adaptive transformer architectures for biomedical natural language "
        "processing. We propose a novel gated attention mechanism that dynamically interpolates between "
        "general and domain-specific representations, enabling effective knowledge transfer from large-scale "
        "pre-training to specialized biomedical tasks. Through a three-phase training procedure encompassing "
        "general pre-training, domain-adaptive pre-training, and task-specific fine-tuning, our model achieves "
        "state-of-the-art results on named entity recognition, relation extraction, and question answering "
        "benchmarks in the biomedical domain. Extensive ablation studies validate the contribution of each "
        "architectural component and training phase."
    )
    doc.add_page_break()

    # ---- Table of Contents placeholder ----
    toc_heading = doc.add_heading("Table of Contents", level=1)
    for i, ch in enumerate(CHAPTERS, 1):
        toc_entry = doc.add_paragraph()
        toc_entry.add_run(f"Chapter {i}: {ch['title']}").bold = True
        for sec_title, _ in ch['sections']:
            doc.add_paragraph(f"    {sec_title}")
    doc.add_page_break()

    # ---- Chapters ----
    for ch_idx, chapter in enumerate(CHAPTERS):
        # Chapter heading (Heading 1 style)
        ch_heading = doc.add_heading(f"Chapter {ch_idx + 1}: {chapter['title']}", level=1)

        for sec_idx, (sec_title, sec_content) in enumerate(chapter['sections']):
            # Section heading (Heading 2)
            doc.add_heading(sec_title, level=2)

            # Split content into multiple paragraphs for realistic length
            doc.add_paragraph(sec_content)

            # Add filler paragraphs to reach ~40 pages total
            doc.add_paragraph(
                f"Further analysis of the {sec_title.lower()} reveals additional complexity that warrants "
                f"careful consideration. The interplay between theoretical foundations and practical "
                f"implementation challenges continues to drive innovation in this rapidly evolving field. "
                f"Researchers have noted that the boundary conditions for optimal performance depend on "
                f"multiple factors including dataset characteristics, model capacity, and computational "
                f"constraints available during the training phase."
            )
            doc.add_paragraph(
                f"In the context of our specific investigation, the {sec_title.lower()} framework provides "
                f"a robust foundation for subsequent experimental validation. The methodology described here "
                f"builds upon established best practices while introducing novel elements that address "
                f"previously unresolved challenges in the domain. Empirical evidence from preliminary "
                f"studies supports the theoretical predictions outlined in earlier sections of this work."
            )

        # Add page break between chapters (except last)
        if ch_idx < len(CHAPTERS) - 1:
            doc.add_page_break()

    # ---- References ----
    doc.add_page_break()
    doc.add_heading("References", level=1)
    references = [
        "Beltagy, I., Lo, K., & Cohan, A. (2019). SciBERT: A pretrained language model for scientific text. EMNLP.",
        "Devlin, J., Chang, M., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. NAACL-HLT.",
        "Hinton, G., Vinyals, O., & Dean, J. (2015). Distilling the knowledge in a neural network. NeurIPS Workshop.",
        "Lan, Z., Chen, M., Goodman, S., Gimpel, K., Sharma, P., & Soricut, R. (2020). ALBERT: A lite BERT for self-supervised learning. ICLR.",
        "Lee, J., Yoon, W., Kim, S., et al. (2020). BioBERT: A pre-trained biomedical language representation model. Bioinformatics, 36(4), 1234-1240.",
        "Liu, Y., Ott, M., Goyal, N., et al. (2019). RoBERTa: A robustly optimized BERT pretraining approach. arXiv:1907.11692.",
        "Pan, S. J., & Yang, Q. (2010). A survey on transfer learning. IEEE Transactions on Knowledge and Data Engineering, 22(10), 1345-1359.",
        "Radford, A., Narasimhan, K., Salimans, T., & Sutskever, I. (2018). Improving language understanding by generative pre-training. OpenAI.",
        "Raffel, C., Shazeer, N., Roberts, A., et al. (2020). Exploring the limits of transfer learning with a unified text-to-text transformer. JMLR, 21(140), 1-67.",
        "Sanh, V., Debut, L., Chaumond, J., & Wolf, T. (2019). DistilBERT, a distilled version of BERT. NeurIPS Workshop.",
        "Vaswani, A., Shazeer, N., Parmar, N., et al. (2017). Attention is all you need. NeurIPS.",
    ]
    for ref in references:
        doc.add_paragraph(ref, style='List Number')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
