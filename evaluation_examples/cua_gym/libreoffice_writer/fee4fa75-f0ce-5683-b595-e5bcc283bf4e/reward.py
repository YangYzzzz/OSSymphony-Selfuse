"""
Reward Script: First page of each chapter has no header, subsequent pages show header with page numbers
Task ID: writer_af_013
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.3): Multiple sections exist (section breaks inserted before chapters)
  - Component 2 (0.3): Sections with titlePg enabled (different first page)
  - Component 3 (0.2): First-page headers are empty (no text, no field codes)
  - Component 4 (0.2): Default headers have page number field codes
"""

import os

from docx import Document

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
RENS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
ns = {'w': WNS}

WORKDIR = '/home/user'
TASK_ID = 'writer_af_013'


def persist_app_state(domain):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Set up a page style where the first page of each chapter has no header
    but all subsequent pages show the header with page numbers.

    Expected golden state:
    - Document has multiple sections (one per chapter boundary)
    - Sections corresponding to chapters have titlePg enabled
    - First-page headers in those sections are empty
    - Default headers in those sections contain page number field codes
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)

    # Component 1: Multiple sections exist (0.3 points)
    # Initial env has only 1 section. Golden should have multiple (at least 3 for chapter breaks).
    try:
        if num_sections >= 3:
            print(f"PASS: Component 1 -- Multiple sections found: {num_sections} (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 -- Expected >= 3 sections, found {num_sections}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: At least 3 sections have titlePg enabled (0.3 points)
    # titlePg enables "different first page" header in a section.
    # Initial env has 0 sections with titlePg. Golden should have >= 3.
    try:
        titlepg_count = 0
        for s in doc.sections:
            sectPr = s._sectPr
            title_pg = sectPr.find('{' + WNS + '}titlePg')
            if title_pg is not None:
                titlepg_count += 1

        if titlepg_count >= 3:
            print(f"PASS: Component 2 -- {titlepg_count} sections with titlePg enabled (0.3 pts)")
            total_score += 0.3
        elif titlepg_count >= 1:
            partial = 0.1 * titlepg_count
            print(f"PARTIAL: Component 2 -- {titlepg_count} sections with titlePg (partial: {partial:.1f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- No sections with titlePg enabled")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: First-page headers are empty in sections with titlePg (0.2 points)
    # The first page of each chapter should have NO header content.
    try:
        sections_with_titlepg = []
        for i, s in enumerate(doc.sections):
            sectPr = s._sectPr
            title_pg = sectPr.find('{' + WNS + '}titlePg')
            if title_pg is not None:
                sections_with_titlepg.append((i, s))

        if len(sections_with_titlepg) == 0:
            print(f"FAIL: Component 3 -- No sections with titlePg to check first-page headers")
        else:
            empty_first_headers = 0
            for idx, s in sections_with_titlepg:
                first_hdr = s.first_page_header
                has_content = False
                if first_hdr and first_hdr.paragraphs:
                    for p in first_hdr.paragraphs:
                        if p.text.strip():
                            has_content = True
                        # Also check for field codes (like PAGE numbers)
                        for el in p._element.findall('.//w:instrText', ns):
                            has_content = True
                if not has_content:
                    empty_first_headers += 1

            if empty_first_headers == len(sections_with_titlepg):
                print(f"PASS: Component 3 -- All {empty_first_headers} first-page headers are empty (0.2 pts)")
                total_score += 0.2
            elif empty_first_headers > 0:
                ratio = empty_first_headers / len(sections_with_titlepg)
                partial = round(0.2 * ratio, 2)
                print(f"PARTIAL: Component 3 -- {empty_first_headers}/{len(sections_with_titlepg)} first-page headers empty ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 -- No first-page headers are empty")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Default headers in sections with titlePg have page number fields (0.2 points)
    # Continuation pages should show header with page numbers.
    try:
        sections_with_titlepg = []
        for i, s in enumerate(doc.sections):
            sectPr = s._sectPr
            title_pg = sectPr.find('{' + WNS + '}titlePg')
            if title_pg is not None:
                sections_with_titlepg.append((i, s))

        if len(sections_with_titlepg) == 0:
            print(f"FAIL: Component 4 -- No sections with titlePg to check default headers")
        else:
            headers_with_pagenum = 0
            for idx, s in sections_with_titlepg:
                default_hdr = s.header
                has_page_field = False
                if default_hdr and default_hdr.paragraphs:
                    for p in default_hdr.paragraphs:
                        for el in p._element.findall('.//w:instrText', ns):
                            if 'PAGE' in el.text.upper():
                                has_page_field = True
                if has_page_field:
                    headers_with_pagenum += 1

            if headers_with_pagenum == len(sections_with_titlepg):
                print(f"PASS: Component 4 -- All {headers_with_pagenum} default headers have page numbers (0.2 pts)")
                total_score += 0.2
            elif headers_with_pagenum > 0:
                ratio = headers_with_pagenum / len(sections_with_titlepg)
                partial = round(0.2 * ratio, 2)
                print(f"PARTIAL: Component 4 -- {headers_with_pagenum}/{len(sections_with_titlepg)} default headers have page numbers ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 -- No default headers have page number fields")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
