"""
Initial Setup: Field Trip Permission Slip - initial state
Task ID: writer_creative_049
Domain: libreoffice_writer

Creates a plain, unformatted permission slip document at /home/user/Desktop/field_trip_permission.docx
All text is 12pt, left-aligned, with no bold/center/underline formatting.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_creative_049'
OUTPUT = f'{WORKDIR}/field_trip_permission.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Remove default styles spacing to keep it simple
    # All paragraphs: 12pt, left-aligned, no special formatting

    lines = [
        'Washington Elementary School',
        '2500 NE Broadway, Portland, OR 97232',
        '',
        'Field Trip Permission Slip',
        '',
        'Dear Parents/Guardians,',
        (
            'We are excited to announce that our class will be taking a field trip to the '
            'Oregon Museum of Science and Industry (OMSI) on March 21, 2026. This is a '
            'wonderful opportunity for students to explore hands-on science exhibits and '
            'deepen their understanding of the natural world. We look forward to a fun and '
            'educational day!'
        ),
        'Departure: 8:30 AM, Return: 2:30 PM, Cost: $15 per student',
        'Students must bring a packed lunch and wear comfortable walking shoes.',
        '',
        'Student Name: _______________',
        'I give permission for my child to attend the OMSI field trip.',
        'Parent/Guardian Signature: _______________ Date: ___________',
        'Emergency Phone: _______________',
    ]

    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        if line:
            run = para.add_run(line)
            run.font.size = Pt(12)
            run.bold = False
            run.italic = False
            run.underline = False

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
