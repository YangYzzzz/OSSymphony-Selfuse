"""
Reward Script: Field Trip Permission Slip Formatting
Task ID: writer_creative_049
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): School name is 16pt, bold, centered
  Component 2 (0.15): School address is 10pt, centered
  Component 3 (0.25): 'Field Trip Permission Slip' is 14pt, bold, underline, centered
  Component 4 (0.15): Trip details line is bold
  Component 5 (0.15): Dashed separator line exists (centered) separating tear-off section
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_creative_049'
FILE_PATH = f'{WORKDIR}/field_trip_permission.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Helper: find paragraph index by text
    def find_para_by_text(doc, search_text):
        for i, para in enumerate(doc.paragraphs):
            if search_text.lower() in para.text.lower():
                return i, para
        return None, None

    # Helper: get effective font size in pt from first run of a paragraph
    def get_para_font_size(para):
        for run in para.runs:
            if run.font.size is not None:
                return run.font.size.pt
        return None

    # Helper: get effective bold from first run of a paragraph
    def get_para_bold(para):
        for run in para.runs:
            return run.bold  # None = inherited, True = bold
        return None

    # Helper: check if any run is underline in paragraph
    def get_para_underline(para):
        for run in para.runs:
            if run.underline:
                return True
        return None

    # Helper: check alignment — CENTER is WD_PARAGRAPH_ALIGNMENT.CENTER (1)
    def is_centered(para):
        return para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Component 1: School name — 16pt, bold, centered (0.30 points) ---
    try:
        idx, school_para = find_para_by_text(doc, 'Washington Elementary School')
        if school_para is not None:
            sz = get_para_font_size(school_para)
            bold = get_para_bold(school_para)
            centered = is_centered(school_para)

            size_ok = (sz is not None and abs(sz - 16.0) < 0.5)
            bold_ok = (bold is True)
            center_ok = centered

            if size_ok and bold_ok and center_ok:
                print(f"PASS: Component 1 — School name is {sz}pt, bold={bold}, centered={centered} (0.30 pts)")
                total_score += 0.30
            else:
                print(f"FAIL: Component 1 — School name: size={sz}pt (expected 16), bold={bold} (expected True), centered={centered} (expected True)")
        else:
            print("FAIL: Component 1 — 'Washington Elementary School' paragraph not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # --- Component 2: School address — 10pt, centered (0.15 points) ---
    try:
        idx, addr_para = find_para_by_text(doc, '2500 NE Broadway')
        if addr_para is not None:
            sz = get_para_font_size(addr_para)
            centered = is_centered(addr_para)

            size_ok = (sz is not None and abs(sz - 10.0) < 0.5)
            center_ok = centered

            if size_ok and center_ok:
                print(f"PASS: Component 2 — School address is {sz}pt, centered={centered} (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 2 — School address: size={sz}pt (expected 10), centered={centered} (expected True)")
        else:
            print("FAIL: Component 2 — School address paragraph not found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # --- Component 3: 'Field Trip Permission Slip' — 14pt, bold, underline, centered (0.25 points) ---
    try:
        idx, title_para = find_para_by_text(doc, 'Field Trip Permission Slip')
        if title_para is not None:
            sz = get_para_font_size(title_para)
            bold = get_para_bold(title_para)
            underline = get_para_underline(title_para)
            centered = is_centered(title_para)

            size_ok = (sz is not None and abs(sz - 14.0) < 0.5)
            bold_ok = (bold is True)
            underline_ok = (underline is True)
            center_ok = centered

            if size_ok and bold_ok and underline_ok and center_ok:
                print(f"PASS: Component 3 — Title is {sz}pt, bold={bold}, underline={underline}, centered={centered} (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — Title: size={sz}pt (exp 14), bold={bold} (exp True), underline={underline} (exp True), centered={centered} (exp True)")
        else:
            print("FAIL: Component 3 — 'Field Trip Permission Slip' paragraph not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # --- Component 4: Trip details line — bold (0.15 points) ---
    try:
        idx, details_para = find_para_by_text(doc, 'Departure:')
        if details_para is not None:
            bold = get_para_bold(details_para)
            bold_ok = (bold is True)
            if bold_ok:
                print(f"PASS: Component 4 — Trip details line is bold (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 4 — Trip details line bold={bold} (expected True)")
        else:
            print("FAIL: Component 4 — Trip details (Departure) paragraph not found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # --- Component 5: Dashed separator line exists (0.15 points) ---
    # Expect a paragraph containing dashes/scissors (- - - or ✂) between the body and the tear-off section
    try:
        sep_count = sum(
            1 for para in doc.paragraphs
            if ('- -' in para.text or '---' in para.text
                or '\u2702' in para.text or '----' in para.text)
        )
        if sep_count > 0:
            print(f"PASS: Component 5 — Dashed separator found ({sep_count} matching paragraph(s)) (0.15 pts)")
            total_score += 0.15
        else:
            print("FAIL: Component 5 — No dashed separator line found in document")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {round(total_score, 2)}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
