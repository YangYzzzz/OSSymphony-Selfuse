"""
Initial Setup: Vowel/Consonant Coloring Task - Academic Abstract
Task ID: osworld_writer_vowel_consonant_coloring_003
Domain: libreoffice_writer

Creates a .docx academic abstract with two paragraphs.
The second paragraph has 40 words with NO custom font colors applied.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_vowel_consonant_coloring_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- First paragraph: Introduction ---
    para1 = doc.add_paragraph()
    run1 = para1.add_run(
        "Climate change represents one of the most pressing scientific and societal challenges "
        "of the twenty-first century. Researchers worldwide have dedicated substantial efforts "
        "to understanding its causes, mechanisms, and far-reaching consequences for both natural "
        "and human systems."
    )
    run1.font.size = Pt(12)

    # --- Second paragraph: Exactly 40 words, NO custom font colors ---
    # Words: An empirical analysis of environmental changes indicates that average global
    # temperatures across major urban regions have elevated by approximately two degrees over
    # the past decade, directly affecting ecosystems, agricultural yields, key biodiversity
    # patterns, and overall community resilience in significant ways.
    para2 = doc.add_paragraph()
    run2 = para2.add_run(
        "An empirical analysis of environmental changes indicates that average global "
        "temperatures across major urban regions have elevated by approximately two degrees over "
        "the past decade, directly affecting ecosystems, agricultural yields, key biodiversity "
        "patterns, and overall community resilience in significant ways."
    )
    run2.font.size = Pt(12)
    # NOTE: No custom font color is applied — run2.font.color.rgb remains None

    # --- Third paragraph: Conclusion ---
    para3 = doc.add_paragraph()
    run3 = para3.add_run(
        "Mitigation strategies and adaptation policies must therefore be grounded in rigorous "
        "empirical evidence and interdisciplinary collaboration to effectively address these "
        "multifaceted challenges and protect vulnerable populations."
    )
    run3.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open file in LibreOffice Writer with DISPLAY=:0
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
