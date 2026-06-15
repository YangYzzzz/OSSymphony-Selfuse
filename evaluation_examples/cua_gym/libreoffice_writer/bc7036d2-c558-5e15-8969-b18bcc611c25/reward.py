"""
Reward Script: Apply red color to vowel-starting words and blue color to consonant-starting
               words in the second paragraph of a document.
Task ID: osworld_writer_vowel_consonant_coloring_003
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): All words in the second paragraph have a color applied (red or blue)
  Component 2 (0.5 pts): Each word's color correctly follows the vowel/consonant rule
                         (red=vowel start, blue=consonant start)
"""

import os
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_vowel_consonant_coloring_003'
VOWELS = set('AEIOUaeiou')


def color_distance(c1, c2):
    """Euclidean distance between two RGB tuples."""
    return sqrt(sum((a - b) ** 2 for a, b in zip(c1, c2)))


def is_red(rgb):
    """Check if an RGBColor is close to pure red (255, 0, 0)."""
    return color_distance(tuple(rgb), (255, 0, 0)) < 50


def is_blue(rgb):
    """Check if an RGBColor is close to pure blue (0, 0, 255)."""
    return color_distance(tuple(rgb), (0, 0, 255)) < 50


def get_word_first_char(word_text):
    """Get the first alphabetic character of a word, stripping leading punctuation."""
    for ch in word_text:
        if ch.isalpha():
            return ch
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.shared import RGBColor
    except ImportError as e:
        print(f"CRITICAL: python-docx not available: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: document must have at least 2 paragraphs
    if len(doc.paragraphs) < 2:
        print(f"CRITICAL: Document has fewer than 2 paragraphs ({len(doc.paragraphs)} found)")
        print("REWARD: 0.0")
        return 0.0

    # Target: second paragraph (index 1)
    para = doc.paragraphs[1]

    # Collect word-runs (non-empty, non-whitespace runs)
    word_runs = [run for run in para.runs if run.text.strip()]

    if not word_runs:
        print("FAIL: Second paragraph has no word-containing runs")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All word-runs have a color applied (red or blue) (0.5 points)
    # This checks that the agent actually applied coloring — in the initial state,
    # there is one uncolored run. After the task, every word-run must have red or blue.
    try:
        colored_count = 0
        uncolored_words = []

        for run in word_runs:
            rgb = None
            try:
                if run.font.color and run.font.color.type is not None:
                    rgb = run.font.color.rgb
            except Exception:
                pass

            if rgb is not None and (is_red(rgb) or is_blue(rgb)):
                colored_count += 1
            else:
                uncolored_words.append(run.text)

        if colored_count == len(word_runs) and len(word_runs) > 0:
            print(f"PASS: Component 1 — All {colored_count} word-runs have red or blue color applied (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — {colored_count}/{len(word_runs)} word-runs have color applied; "
                  f"uncolored: {uncolored_words[:5]}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Colors are correct per vowel/consonant rule (0.5 points)
    # Each word-run's color must match the rule:
    #   starts with a vowel (A,E,I,O,U) → red (255, 0, 0)
    #   starts with a consonant          → blue (0, 0, 255)
    try:
        correctly_colored = 0
        incorrectly_colored = []

        for run in word_runs:
            rgb = None
            try:
                if run.font.color and run.font.color.type is not None:
                    rgb = run.font.color.rgb
            except Exception:
                pass

            first_char = get_word_first_char(run.text)
            if first_char is None:
                # Punctuation-only run, skip
                continue

            if first_char in VOWELS:
                expected_color = 'red'
                color_correct = rgb is not None and is_red(rgb)
            else:
                expected_color = 'blue'
                color_correct = rgb is not None and is_blue(rgb)

            if color_correct:
                correctly_colored += 1
            else:
                actual_hex = str(rgb) if rgb else 'None'
                incorrectly_colored.append(
                    f"'{run.text}' starts_with='{first_char}' expected={expected_color} got={actual_hex}"
                )

        # Count all words with an alphabetic first char
        total_alpha_words = sum(
            1 for run in word_runs if get_word_first_char(run.text) is not None
        )

        if total_alpha_words > 0 and correctly_colored == total_alpha_words:
            print(f"PASS: Component 2 — All {correctly_colored} words correctly colored "
                  f"per vowel/consonant rule (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 2 — {correctly_colored}/{total_alpha_words} words correctly colored; "
                  f"first mismatches: {incorrectly_colored[:3]}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the environment
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
