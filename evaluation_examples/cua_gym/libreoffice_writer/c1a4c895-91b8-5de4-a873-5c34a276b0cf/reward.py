"""
Reward Script: Workplace Health and Safety Policy Document
Task ID: writer_hr_092
Domain: libreoffice_writer
Scoring:
  Component 1: Structured sections with Heading 1 headings (0.20)
  Component 2: PPE Requirements table (0.15)
  Component 3: Incident Reporting Form table (0.10)
  Component 4: Emergency Numbers table (0.10)
  Component 5: SDS Reference table (0.10)
  Component 6: Training Requirements Matrix table (0.10)
  Component 7: Warning symbols before critical sections (0.10)
  Component 8: Accident Investigation Checklist items (0.05)
  Component 9: Ergonomic Assessment Checklist items (0.05)
  Component 10: Numbered procedure lists (0.05)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_092'


def find_table_by_header(tables, keywords):
    """Find a table whose header row contains any of the given keywords (case-insensitive)."""
    for table in tables:
        if len(table.rows) < 2:
            continue
        header_text = ' '.join(cell.text.strip().lower() for cell in table.rows[0].cells)
        if any(kw.lower() in header_text for kw in keywords):
            return table
    return None


def find_table_by_cell_content(tables, keywords):
    """Find a table where any cell in the first column contains any keyword."""
    for table in tables:
        if len(table.rows) < 2:
            continue
        col0_text = ' '.join(row.cells[0].text.strip().lower() for row in table.rows)
        if any(kw.lower() in col0_text for kw in keywords):
            return table
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_text = '\n'.join(p.text for p in doc.paragraphs)
    tables = doc.tables

    # Component 1: Structured sections with Heading 1 headings (0.20 points)
    # Initial doc has NO headings - all Normal style. Golden has 10 Heading 1 sections.
    try:
        heading1_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Heading 1']
        heading1_count = len(heading1_paras)
        if heading1_count >= 8:
            print(f"PASS: Component 1 - Found {heading1_count} Heading 1 sections (>=8 required) (0.20 pts)")
            total_score += 0.20
        elif heading1_count >= 5:
            partial = 0.10
            print(f"PARTIAL: Component 1 - Found {heading1_count} Heading 1 sections (5-7, partial credit) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 - Found {heading1_count} Heading 1 sections, need >=8")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: PPE Requirements table (0.15 points)
    # Must have columns related to work area and PPE, with at least 5 data rows
    try:
        ppe_table = find_table_by_header(tables, ['work area', 'ppe', 'required ppe'])
        if ppe_table and len(ppe_table.rows) >= 5 and len(ppe_table.columns) >= 3:
            print(f"PASS: Component 2 - PPE table found with {len(ppe_table.rows)} rows, {len(ppe_table.columns)} cols (0.15 pts)")
            total_score += 0.15
        elif ppe_table:
            partial = 0.07
            print(f"PARTIAL: Component 2 - PPE table found but only {len(ppe_table.rows)} rows / {len(ppe_table.columns)} cols ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 - No PPE requirements table found")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Incident Reporting Form table (0.10 points)
    # A form-style table with fields like Date of Incident, Location, Description
    try:
        incident_table = find_table_by_cell_content(tables, ['date of incident', 'description of incident', 'location'])
        if incident_table and len(incident_table.rows) >= 5:
            print(f"PASS: Component 3 - Incident reporting form table found with {len(incident_table.rows)} rows (0.10 pts)")
            total_score += 0.10
        elif incident_table:
            partial = 0.05
            print(f"PARTIAL: Component 3 - Incident form table found but only {len(incident_table.rows)} rows ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 - No incident reporting form table found")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Emergency Numbers table (0.10 points)
    # Table with contact info including phone/extension numbers
    try:
        emergency_table = find_table_by_header(tables, ['contact', 'phone', 'extension', 'emergency'])
        if emergency_table is None:
            # Also check for table with 911 or poison control in it
            for table in tables:
                for row in table.rows:
                    row_text = ' '.join(cell.text.lower() for cell in row.cells)
                    if '911' in row_text or 'poison control' in row_text or 'first aid' in row_text.lower():
                        emergency_table = table
                        break
                if emergency_table:
                    break

        if emergency_table and len(emergency_table.rows) >= 4:
            print(f"PASS: Component 4 - Emergency numbers table found with {len(emergency_table.rows)} rows (0.10 pts)")
            total_score += 0.10
        elif emergency_table:
            partial = 0.05
            print(f"PARTIAL: Component 4 - Emergency numbers table found but only {len(emergency_table.rows)} rows ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 - No emergency contact numbers table found")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: SDS Reference table (0.10 points)
    # Table with chemical names, CAS numbers, hazard classes
    try:
        sds_table = find_table_by_header(tables, ['chemical name', 'cas number', 'hazard class', 'sds'])
        if sds_table and len(sds_table.rows) >= 4 and len(sds_table.columns) >= 4:
            print(f"PASS: Component 5 - SDS reference table found with {len(sds_table.rows)} rows, {len(sds_table.columns)} cols (0.10 pts)")
            total_score += 0.10
        elif sds_table:
            partial = 0.05
            print(f"PARTIAL: Component 5 - SDS table found but small ({len(sds_table.rows)} rows, {len(sds_table.columns)} cols) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 - No SDS reference table found")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: Training Requirements Matrix table (0.10 points)
    # A roles x trainings matrix with at least 5 rows and 4 columns
    try:
        training_table = find_table_by_header(tables, ['training', 'role', 'all employees', 'supervisor'])
        if training_table and len(training_table.rows) >= 5 and len(training_table.columns) >= 4:
            print(f"PASS: Component 6 - Training matrix found with {len(training_table.rows)} rows, {len(training_table.columns)} cols (0.10 pts)")
            total_score += 0.10
        elif training_table:
            partial = 0.05
            print(f"PARTIAL: Component 6 - Training matrix found but small ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 6 - No training requirements matrix found")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    # Component 7: Warning symbols before critical sections (0.10 points)
    # At least 3 headings/paragraphs should contain warning symbols (unicode warning sign or similar)
    try:
        warning_chars = ['\u26a0', '\u2622', '\u2623', '\u26d4', '\u2620', '\u26a1']
        warning_count = 0
        for p in doc.paragraphs:
            if p.style and p.style.name.startswith('Heading'):
                if any(wc in p.text for wc in warning_chars):
                    warning_count += 1

        if warning_count >= 3:
            print(f"PASS: Component 7 - Found {warning_count} headings with warning symbols (>=3 required) (0.10 pts)")
            total_score += 0.10
        elif warning_count >= 1:
            partial = 0.05
            print(f"PARTIAL: Component 7 - Found {warning_count} warning symbols (need >=3) ({partial} pts)")
            total_score += partial
        else:
            # Also check non-heading paragraphs for warning symbols near section starts
            any_warning = sum(1 for p in doc.paragraphs if any(wc in p.text for wc in warning_chars))
            if any_warning >= 3:
                print(f"PASS: Component 7 - Found {any_warning} paragraphs with warning symbols (0.10 pts)")
                total_score += 0.10
            elif any_warning >= 1:
                partial = 0.05
                print(f"PARTIAL: Component 7 - Found {any_warning} paragraphs with warning symbols ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 7 - No warning symbols found in document")
    except Exception as e:
        print(f"ERROR: Component 7 - {e}")

    # Component 8: Accident Investigation Checklist items (0.05 points)
    # Bullet/checkbox items related to accident investigation
    try:
        # Look for checklist items (checkbox character or List Bullet) near accident investigation section
        checklist_items = []
        in_accident_section = False
        for p in doc.paragraphs:
            text_lower = p.text.lower()
            if 'accident investigation' in text_lower and p.style and p.style.name.startswith('Heading'):
                in_accident_section = True
                continue
            if in_accident_section and p.style and p.style.name.startswith('Heading'):
                in_accident_section = False
                continue
            if in_accident_section and p.text.strip():
                if '\u2610' in p.text or (p.style and 'Bullet' in p.style.name) or (p.style and 'List' in p.style.name):
                    checklist_items.append(p.text.strip())

        if len(checklist_items) >= 8:
            print(f"PASS: Component 8 - Found {len(checklist_items)} accident investigation checklist items (>=8 required) (0.05 pts)")
            total_score += 0.05
        elif len(checklist_items) >= 4:
            partial = 0.025
            print(f"PARTIAL: Component 8 - Found {len(checklist_items)} checklist items (need >=8) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 8 - Found {len(checklist_items)} accident investigation checklist items, need >=8")
    except Exception as e:
        print(f"ERROR: Component 8 - {e}")

    # Component 9: Ergonomic Assessment Checklist items (0.05 points)
    # Bullet/checkbox items related to ergonomic assessment
    try:
        ergo_items = []
        in_ergo_section = False
        for p in doc.paragraphs:
            text_lower = p.text.lower()
            if 'ergonomic' in text_lower and p.style and p.style.name.startswith('Heading'):
                in_ergo_section = True
                continue
            if in_ergo_section and p.style and p.style.name.startswith('Heading'):
                in_ergo_section = False
                continue
            if in_ergo_section and p.text.strip():
                if '\u2610' in p.text or (p.style and 'Bullet' in p.style.name) or (p.style and 'List' in p.style.name):
                    ergo_items.append(p.text.strip())

        if len(ergo_items) >= 6:
            print(f"PASS: Component 9 - Found {len(ergo_items)} ergonomic checklist items (>=6 required) (0.05 pts)")
            total_score += 0.05
        elif len(ergo_items) >= 3:
            partial = 0.025
            print(f"PARTIAL: Component 9 - Found {len(ergo_items)} ergonomic items (need >=6) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 9 - Found {len(ergo_items)} ergonomic checklist items, need >=6")
    except Exception as e:
        print(f"ERROR: Component 9 - {e}")

    # Component 10: Numbered procedure lists (0.05 points)
    # Numbered list items for procedures like hazard ID steps, fire evacuation steps
    try:
        numbered_items = [p for p in doc.paragraphs if p.style and p.style.name == 'List Number']
        if len(numbered_items) >= 8:
            print(f"PASS: Component 10 - Found {len(numbered_items)} numbered list items (>=8 required) (0.05 pts)")
            total_score += 0.05
        elif len(numbered_items) >= 4:
            partial = 0.025
            print(f"PARTIAL: Component 10 - Found {len(numbered_items)} numbered items (need >=8) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 10 - Found {len(numbered_items)} numbered list items, need >=8")
    except Exception as e:
        print(f"ERROR: Component 10 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
