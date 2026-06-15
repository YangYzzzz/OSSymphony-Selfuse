"""
Initial Setup: Workplace Health and Safety Policy - unstructured draft
Task ID: writer_hr_092
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_092'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Workplace_Health_Safety_Policy', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle / date line
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sub.add_run('Effective Date: January 15, 2025')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # Unstructured safety content - no sections, no tables, no checklists
    doc.add_paragraph(
        'This document outlines the workplace health and safety policies for '
        'Meridian Industrial Solutions. All employees are expected to comply with '
        'the guidelines described herein. Safety is a shared responsibility and '
        'every team member plays a vital role in maintaining a safe working environment.'
    )

    doc.add_paragraph(
        'Employees must follow general safety rules at all times. This includes '
        'wearing appropriate clothing, keeping work areas clean and free of obstructions, '
        'reporting any unsafe conditions immediately to supervisors, and never operating '
        'equipment without proper training. Horseplay and the use of controlled substances '
        'on company premises are strictly prohibited.'
    )

    doc.add_paragraph(
        'Hazard identification is a continuous process. Workers should be aware of '
        'potential hazards in their work areas including slip and trip hazards, electrical '
        'risks, chemical exposure, moving machinery, and ergonomic strain. All identified '
        'hazards should be reported using the internal reporting system. Regular workplace '
        'inspections are conducted by the safety committee on a quarterly basis.'
    )

    doc.add_paragraph(
        'Personal protective equipment is required in many areas of the facility. '
        'Hard hats are mandatory in the warehouse and loading dock. Safety glasses '
        'must be worn in the laboratory and manufacturing floor. Chemical-resistant '
        'gloves are required when handling solvents in the chemical storage area. '
        'Steel-toed boots are required in all production areas. Hearing protection '
        'must be used in areas where noise levels exceed 85 decibels.'
    )

    doc.add_paragraph(
        'When an incident occurs, it must be reported within 24 hours. The report '
        'should include the date and time of the incident, the location, the names '
        'of all persons involved, a description of what happened, any injuries sustained, '
        'and the names of witnesses. The safety manager, currently Rachel Torres, must '
        'be notified directly for any incident resulting in injury or property damage '
        'exceeding $500.'
    )

    doc.add_paragraph(
        'Accident investigations should be conducted promptly to determine root causes. '
        'The investigation team should examine the scene, interview witnesses, review '
        'relevant procedures, check equipment maintenance records, and assess whether '
        'training was adequate. Corrective actions must be documented and tracked to '
        'completion. The VP of Operations, David Nakamura, reviews all investigation '
        'reports monthly.'
    )

    doc.add_paragraph(
        'First aid kits are located at 12 stations throughout the facility. Trained '
        'first aiders include Jennifer Walsh (Ext. 4201), Marcus Chen (Ext. 4305), '
        'and Anika Patel (Ext. 4118). For medical emergencies call 911 immediately. '
        'The nearest hospital is St. Elizabeth Regional Medical Center at 2400 Oak '
        'Boulevard, approximately 8 minutes by ambulance. The Poison Control Center '
        'can be reached at 1-800-222-1222.'
    )

    doc.add_paragraph(
        'In case of fire, activate the nearest fire alarm pull station and evacuate '
        'using the designated exit routes posted throughout the building. Do not use '
        'elevators. Assembly points are located in Parking Lot A (north side) and '
        'Parking Lot C (south side). Floor wardens are responsible for ensuring all '
        'personnel have evacuated their assigned zones. Fire drills are conducted '
        'semi-annually.'
    )

    doc.add_paragraph(
        'Chemical handling must follow Safety Data Sheet guidelines. All chemicals '
        'must be properly labeled and stored in designated cabinets. Incompatible '
        'chemicals must be stored separately. Common chemicals used on site include '
        'acetone, isopropyl alcohol, sodium hydroxide, hydrochloric acid, and toluene. '
        'SDS binders are maintained at each chemical storage location and digital copies '
        'are available on the company intranet.'
    )

    doc.add_paragraph(
        'Ergonomic assessments should be performed for all workstations. Key factors '
        'include monitor height and distance, chair height and lumbar support, keyboard '
        'and mouse positioning, lighting adequacy, and frequency of breaks. Employees '
        'experiencing discomfort should request an ergonomic evaluation through HR. '
        'Standing desks are available upon request.'
    )

    doc.add_paragraph(
        'Training requirements vary by role. New hires must complete general safety '
        'orientation within the first week. Forklift operators need certification '
        'renewed every three years. Laboratory personnel require chemical safety '
        'training annually. All employees must complete fire safety training yearly. '
        'Supervisors and managers must complete OSHA 30-hour training. First aid and '
        'CPR certification is required for designated responders every two years.'
    )

    doc.add_paragraph(
        'This policy is reviewed annually by the Safety Committee chaired by '
        'Chief Safety Officer Linda Bergstrom. Questions or suggestions regarding '
        'this policy should be directed to the Safety Department at safety@meridian-industrial.com '
        'or extension 4500.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
