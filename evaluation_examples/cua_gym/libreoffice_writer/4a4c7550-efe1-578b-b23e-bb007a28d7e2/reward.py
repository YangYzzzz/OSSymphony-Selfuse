"""
Reward Script: Strategic Workforce Planning Document with Complex Tables
Task ID: writer_hr_078
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25) - At least 5 tables exist in the document
  Component 2 (0.20) - "List of Tables" section with >= 5 entries
  Component 3 (0.15) - Table captions (>= 5 "Table N:" paragraphs near table content)
  Component 4 (0.15) - Demographics table: departments as rows, job levels as columns, with totals
  Component 5 (0.10) - Retirement risk table: departments with 1/3/5-year windows
  Component 6 (0.10) - Skills gap analysis table: competencies with current, required, gap scores
  Component 7 (0.05) - Hiring plan tables: at least 3 quarterly hiring tables
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_078'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_tables = len(doc.tables)
    paragraphs = doc.paragraphs

    # Component 1: At least 5 tables exist (0.25 points)
    # Initial env has 0 tables; golden has 7. This checks the core task requirement.
    try:
        if num_tables >= 5:
            print(f"PASS: Component 1 — {num_tables} tables found (>= 5 required) (0.25 pts)")
            total_score += 0.25
        elif num_tables >= 3:
            partial = 0.15
            print(f"PARTIAL: Component 1 — {num_tables} tables found (>= 5 ideal, >= 3 partial) ({partial} pts)")
            total_score += partial
        elif num_tables >= 1:
            partial = 0.05
            print(f"PARTIAL: Component 1 — {num_tables} tables found (>= 5 ideal) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: "List of Tables" section with >= 5 entries (0.20 points)
    # Initial env has no List of Tables. Golden has a heading + 5 table references.
    try:
        lot_heading_idx = None
        for i, p in enumerate(paragraphs):
            text_lower = p.text.strip().lower()
            if 'list of tables' in text_lower:
                style = p.style.name if p.style else ''
                # Accept as heading or bold paragraph
                if 'heading' in style.lower() or 'title' in style.lower() or p.text.strip():
                    lot_heading_idx = i
                    break

        if lot_heading_idx is not None:
            # Count "Table N:" entries after the heading until next heading or blank gap
            lot_entries = 0
            for j in range(lot_heading_idx + 1, min(lot_heading_idx + 15, len(paragraphs))):
                p_text = paragraphs[j].text.strip()
                if not p_text:
                    # Empty paragraph signals end of list
                    break
                style_name = paragraphs[j].style.name if paragraphs[j].style else ''
                if 'heading' in style_name.lower():
                    break
                if re.match(r'Table\s+\d+\s*[:.]', p_text, re.IGNORECASE):
                    lot_entries += 1

            if lot_entries >= 5:
                print(f"PASS: Component 2 — List of Tables found with {lot_entries} entries (0.20 pts)")
                total_score += 0.20
            elif lot_entries >= 3:
                partial = 0.10
                print(f"PARTIAL: Component 2 — List of Tables found with {lot_entries} entries ({partial} pts)")
                total_score += partial
            elif lot_entries >= 1:
                partial = 0.05
                print(f"PARTIAL: Component 2 — List of Tables found with {lot_entries} entries ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — List of Tables heading found but no 'Table N:' entries after it")
        else:
            print(f"FAIL: Component 2 — No 'List of Tables' section found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Table captions near tables (>= 5) (0.15 points)
    # Captions are paragraphs with "Table N:" pattern that appear near actual table content
    # (not just in the List of Tables). Initial has no captions. Golden has 5.
    try:
        # Collect caption paragraphs that are NOT in the List of Tables section
        caption_count = 0
        past_lot = False
        for i, p in enumerate(paragraphs):
            text = p.text.strip()
            # Skip List of Tables section
            if 'list of tables' in text.lower():
                past_lot = False
                continue
            style_name = p.style.name if p.style else ''
            if 'heading' in style_name.lower() and 'list of tables' not in text.lower():
                past_lot = True
            if past_lot and re.match(r'Table\s+\d+\s*[:.]', text, re.IGNORECASE):
                caption_count += 1

        if caption_count >= 5:
            print(f"PASS: Component 3 — {caption_count} table captions found (>= 5) (0.15 pts)")
            total_score += 0.15
        elif caption_count >= 3:
            partial = 0.08
            print(f"PARTIAL: Component 3 — {caption_count} table captions found ({partial} pts)")
            total_score += partial
        elif caption_count >= 1:
            partial = 0.03
            print(f"PARTIAL: Component 3 — {caption_count} table captions found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No table captions found outside List of Tables")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Demographics table structure (0.15 points)
    # Must have departments as rows with job level columns (Entry, Mid, Senior, etc.)
    # Initial has this data as narrative text, not in a table.
    try:
        demo_found = False
        if num_tables >= 1:
            for table in doc.tables:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                # Check for department + level columns pattern
                has_dept = any('department' in h for h in headers)
                has_levels = (
                    any('entry' in h for h in headers) or
                    any('mid' in h for h in headers) or
                    any('senior' in h for h in headers)
                )
                has_total = any('total' in h for h in headers)

                if has_dept and has_levels:
                    demo_found = True
                    # Check it has multiple department rows
                    dept_rows = 0
                    for row in table.rows[1:]:
                        cell_text = row.cells[0].text.strip()
                        if cell_text and cell_text.lower() != 'total':
                            dept_rows += 1
                    if dept_rows >= 5 and has_total:
                        print(f"PASS: Component 4 — Demographics table found: {dept_rows} departments, has total column (0.15 pts)")
                        total_score += 0.15
                    elif dept_rows >= 3:
                        partial = 0.08
                        print(f"PARTIAL: Component 4 — Demographics table found: {dept_rows} departments ({partial} pts)")
                        total_score += partial
                    else:
                        partial = 0.05
                        print(f"PARTIAL: Component 4 — Demographics table found but only {dept_rows} department rows ({partial} pts)")
                        total_score += partial
                    break

        if not demo_found:
            print(f"FAIL: Component 4 — No demographics table with department/level structure found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Retirement risk table (0.10 points)
    # Must have departments with 1-year, 3-year, 5-year retirement windows
    # Initial has this as narrative text only.
    try:
        retire_found = False
        if num_tables >= 2:
            for table in doc.tables:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                header_text = ' '.join(headers)
                has_year_windows = (
                    ('1 year' in header_text or '1year' in header_text or 'within 1' in header_text) and
                    ('3 year' in header_text or '3year' in header_text or 'within 3' in header_text) and
                    ('5 year' in header_text or '5year' in header_text or 'within 5' in header_text)
                )
                has_risk = any('risk' in h for h in headers) or any('rating' in h for h in headers)

                if has_year_windows:
                    retire_found = True
                    row_count = len(table.rows) - 1  # exclude header
                    if row_count >= 5 and has_risk:
                        print(f"PASS: Component 5 — Retirement risk table found: {row_count} rows, has risk rating (0.10 pts)")
                        total_score += 0.10
                    elif row_count >= 3:
                        partial = 0.05
                        print(f"PARTIAL: Component 5 — Retirement risk table found: {row_count} rows ({partial} pts)")
                        total_score += partial
                    else:
                        partial = 0.03
                        print(f"PARTIAL: Component 5 — Retirement risk table found but only {row_count} rows ({partial} pts)")
                        total_score += partial
                    break

        if not retire_found:
            print(f"FAIL: Component 5 — No retirement risk table with 1/3/5-year windows found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Skills gap analysis table (0.10 points)
    # Must have competencies with current level, required level, and gap score columns
    # Initial has this as narrative text only.
    try:
        skills_found = False
        if num_tables >= 3:
            for table in doc.tables:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                header_text = ' '.join(headers)
                has_current = any('current' in h for h in headers)
                has_required = any('required' in h for h in headers)
                has_gap = any('gap' in h for h in headers)

                if has_current and has_required and has_gap:
                    skills_found = True
                    row_count = len(table.rows) - 1
                    if row_count >= 5:
                        print(f"PASS: Component 6 — Skills gap table found: {row_count} competency rows (0.10 pts)")
                        total_score += 0.10
                    elif row_count >= 3:
                        partial = 0.05
                        print(f"PARTIAL: Component 6 — Skills gap table found: {row_count} rows ({partial} pts)")
                        total_score += partial
                    else:
                        partial = 0.03
                        print(f"PARTIAL: Component 6 — Skills gap table found but only {row_count} rows ({partial} pts)")
                        total_score += partial
                    break

        if not skills_found:
            print(f"FAIL: Component 6 — No skills gap analysis table found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Hiring plan tables — at least 3 quarterly hiring tables (0.05 points)
    # These tables have Q1, Q2, Q3, Q4, Annual Total columns.
    # Initial has hiring data as narrative text only.
    try:
        hiring_table_count = 0
        if num_tables >= 4:
            for table in doc.tables:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                has_quarters = (
                    any('q1' in h for h in headers) and
                    any('q2' in h for h in headers) and
                    any('q3' in h for h in headers) and
                    any('q4' in h for h in headers)
                )
                if has_quarters:
                    hiring_table_count += 1

        if hiring_table_count >= 3:
            print(f"PASS: Component 7 — {hiring_table_count} quarterly hiring tables found (>= 3) (0.05 pts)")
            total_score += 0.05
        elif hiring_table_count >= 1:
            partial = 0.02
            print(f"PARTIAL: Component 7 — {hiring_table_count} quarterly hiring tables found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 7 — No quarterly hiring plan tables found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
