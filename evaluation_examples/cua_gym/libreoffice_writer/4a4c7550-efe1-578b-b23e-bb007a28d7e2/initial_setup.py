"""
Initial Setup: Strategic Workforce Planning Document with raw data (no tables)
Task ID: writer_hr_078
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_078'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Document Title --
    title = doc.add_heading('Strategic Workforce Plan 2026-2028', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by Human Resources Division\nEffective Date: January 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()  # spacer

    # ===== EXECUTIVE SUMMARY =====
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This Strategic Workforce Plan outlines the talent management strategy for Meridian Technologies '
        'for the period 2026 through 2028. The plan addresses critical workforce challenges including '
        'an aging workforce, emerging skills gaps in cloud computing and AI/ML disciplines, and '
        'projected retirement waves across key departments. Our analysis covers five core dimensions: '
        'current workforce demographics, retirement risk assessment, succession planning for critical '
        'roles, competency gap analysis, and a phased hiring plan with quarterly milestones.'
    )
    doc.add_paragraph(
        'Key findings indicate that 18% of our senior technical staff will be eligible for retirement '
        'within the next three years, with the Engineering and Research divisions facing the most acute '
        'risk. The succession pipeline currently covers only 62% of mission-critical positions at the '
        'ready-now level. Additionally, competency assessments reveal significant gaps in data science, '
        'cybersecurity, and DevOps practices that must be addressed through both external hiring and '
        'internal upskilling programs.'
    )

    # ===== SECTION 1: CURRENT WORKFORCE DEMOGRAPHICS =====
    doc.add_heading('1. Current Workforce Demographics', level=1)
    doc.add_paragraph(
        'As of Q4 2025, Meridian Technologies employs 1,247 full-time employees across six departments. '
        'The following data summarizes headcount distribution by department, job level, and age band.'
    )

    doc.add_heading('Departmental Headcount by Job Level', level=2)
    doc.add_paragraph(
        'Engineering department: Entry Level 45, Mid Level 78, Senior Level 52, '
        'Principal/Lead 18, Director+ 6, Total 199.'
    )
    doc.add_paragraph(
        'Product Management: Entry Level 12, Mid Level 28, Senior Level 22, '
        'Principal/Lead 8, Director+ 4, Total 74.'
    )
    doc.add_paragraph(
        'Sales & Marketing: Entry Level 38, Mid Level 64, Senior Level 41, '
        'Principal/Lead 14, Director+ 7, Total 164.'
    )
    doc.add_paragraph(
        'Research & Development: Entry Level 22, Mid Level 45, Senior Level 38, '
        'Principal/Lead 15, Director+ 5, Total 125.'
    )
    doc.add_paragraph(
        'Operations: Entry Level 85, Mid Level 142, Senior Level 68, '
        'Principal/Lead 12, Director+ 4, Total 311.'
    )
    doc.add_paragraph(
        'Human Resources & Admin: Entry Level 28, Mid Level 52, Senior Level 34, '
        'Principal/Lead 6, Director+ 3, Total 123.'
    )
    doc.add_paragraph(
        'Finance & Legal: Entry Level 32, Mid Level 58, Senior Level 42, '
        'Principal/Lead 10, Director+ 5, Total 147.'
    )
    doc.add_paragraph(
        'IT Infrastructure: Entry Level 18, Mid Level 35, Senior Level 22, '
        'Principal/Lead 8, Director+ 3, Total 86.'
    )
    doc.add_paragraph(
        'Customer Success: Entry Level 20, Mid Level 38, Senior Level 24, '
        'Principal/Lead 5, Director+ 2, Total 89.'
    )

    doc.add_heading('Age Band Distribution', level=2)
    doc.add_paragraph(
        'Across the organization, the age distribution is as follows: Under 30 years 284 employees (22.8%), '
        '30-39 years 398 employees (31.9%), 40-49 years 312 employees (25.0%), 50-59 years 187 employees '
        '(15.0%), 60 and over 66 employees (5.3%). The 50+ cohort represents 20.3% of total headcount '
        'and is concentrated in senior technical and management roles.'
    )

    # ===== SECTION 2: RETIREMENT RISK ASSESSMENT =====
    doc.add_heading('2. Projected Retirement Risk Assessment', level=1)
    doc.add_paragraph(
        'The retirement risk analysis projects voluntary and mandatory separation due to age-based '
        'retirement eligibility. Employees are categorized into three risk windows based on proximity '
        'to retirement age (65 years).'
    )

    doc.add_heading('Retirement Eligibility by Department', level=2)
    doc.add_paragraph(
        'Engineering: Within 1 year 4 employees, within 3 years 12 employees, within 5 years 23 employees. '
        'Cumulative risk rating: High.'
    )
    doc.add_paragraph(
        'Product Management: Within 1 year 1 employee, within 3 years 4 employees, within 5 years 8 employees. '
        'Cumulative risk rating: Medium.'
    )
    doc.add_paragraph(
        'Sales & Marketing: Within 1 year 3 employees, within 3 years 8 employees, within 5 years 15 employees. '
        'Cumulative risk rating: Medium.'
    )
    doc.add_paragraph(
        'Research & Development: Within 1 year 5 employees, within 3 years 14 employees, within 5 years 22 employees. '
        'Cumulative risk rating: Critical.'
    )
    doc.add_paragraph(
        'Operations: Within 1 year 6 employees, within 3 years 15 employees, within 5 years 28 employees. '
        'Cumulative risk rating: High.'
    )
    doc.add_paragraph(
        'Human Resources & Admin: Within 1 year 2 employees, within 3 years 5 employees, within 5 years 9 employees. '
        'Cumulative risk rating: Low.'
    )
    doc.add_paragraph(
        'Finance & Legal: Within 1 year 3 employees, within 3 years 9 employees, within 5 years 16 employees. '
        'Cumulative risk rating: Medium.'
    )
    doc.add_paragraph(
        'IT Infrastructure: Within 1 year 2 employees, within 3 years 6 employees, within 5 years 11 employees. '
        'Cumulative risk rating: Medium.'
    )
    doc.add_paragraph(
        'Customer Success: Within 1 year 1 employee, within 3 years 3 employees, within 5 years 6 employees. '
        'Cumulative risk rating: Low.'
    )
    doc.add_paragraph(
        'Organization totals: Within 1 year 27 employees, within 3 years 76 employees, within 5 years 138 employees.'
    )

    # ===== SECTION 3: SUCCESSION PLANNING =====
    doc.add_heading('3. Succession Planning Matrix', level=1)
    doc.add_paragraph(
        'The succession planning matrix evaluates readiness levels for each critical role. '
        'Readiness categories are: Ready Now (can assume role within 0-6 months), Ready Soon '
        '(6-18 months with targeted development), Developing (18-36 months, requires significant '
        'investment), and No Candidate Identified.'
    )

    doc.add_heading('Critical Role Succession Status', level=2)
    doc.add_paragraph(
        'Chief Technology Officer (CTO): Ready Now candidate is Priya Sharma (VP Engineering), '
        'Ready Soon candidate is David Kim (Sr. Director Architecture), Developing candidate is '
        'Elena Vasquez (Principal Engineer). Risk Level: Low.'
    )
    doc.add_paragraph(
        'VP of Engineering: Ready Now candidate is Marcus Chen (Sr. Director Platform), '
        'Ready Soon candidate is Aisha Patel (Director Backend Services). '
        'No Developing candidate identified. Risk Level: Medium.'
    )
    doc.add_paragraph(
        'Director of Data Science: No Ready Now candidate. Ready Soon candidate is '
        'James Liu (Sr. Data Scientist). Developing candidate is Fatima Al-Rahman '
        '(Data Scientist II). Risk Level: High.'
    )
    doc.add_paragraph(
        'Head of Cybersecurity: No Ready Now candidate. No Ready Soon candidate. '
        'Developing candidate is Robert Okonkwo (Security Analyst III). Risk Level: Critical.'
    )
    doc.add_paragraph(
        'VP of Sales: Ready Now candidate is Jennifer Brooks (Regional Director West), '
        'Ready Soon candidate is Thomas Wright (Regional Director East), '
        'Developing candidate is Sofia Hernandez (Sr. Account Executive). Risk Level: Low.'
    )
    doc.add_paragraph(
        'Chief Financial Officer (CFO): Ready Now candidate is William Park (VP Finance), '
        'No Ready Soon candidate. Developing candidate is Rachel Goldstein (Controller). '
        'Risk Level: Medium.'
    )
    doc.add_paragraph(
        'Director of Product: No Ready Now candidate. Ready Soon candidate is '
        'Kai Nakamura (Sr. Product Manager). Developing candidate is Laura Mitchell '
        '(Product Manager II). Risk Level: High.'
    )
    doc.add_paragraph(
        'Head of Cloud Infrastructure: Ready Now candidate is Ahmad Hassan (Sr. Cloud Architect), '
        'No Ready Soon candidate. Developing candidate is Nicole Dubois (Cloud Engineer III). '
        'Risk Level: Medium.'
    )

    # ===== SECTION 4: SKILLS GAP ANALYSIS =====
    doc.add_heading('4. Skills Gap Analysis', level=1)
    doc.add_paragraph(
        'A competency assessment was conducted across all technical and managerial roles. '
        'Each competency is rated on a 1-5 scale where 1 = Novice and 5 = Expert. '
        'The gap score represents the difference between the required proficiency level '
        'and the current average proficiency across the relevant workforce segment.'
    )

    doc.add_heading('Competency Assessment Results', level=2)
    doc.add_paragraph(
        'Cloud Architecture (AWS/Azure/GCP): Current average proficiency 2.8, '
        'required proficiency 4.2, gap score -1.4. Priority: Critical.'
    )
    doc.add_paragraph(
        'Machine Learning / AI: Current average proficiency 2.1, '
        'required proficiency 4.0, gap score -1.9. Priority: Critical.'
    )
    doc.add_paragraph(
        'Cybersecurity & Compliance: Current average proficiency 2.5, '
        'required proficiency 4.5, gap score -2.0. Priority: Critical.'
    )
    doc.add_paragraph(
        'DevOps & CI/CD Practices: Current average proficiency 3.0, '
        'required proficiency 4.0, gap score -1.0. Priority: High.'
    )
    doc.add_paragraph(
        'Data Engineering & Analytics: Current average proficiency 2.6, '
        'required proficiency 3.8, gap score -1.2. Priority: High.'
    )
    doc.add_paragraph(
        'Agile/Scrum Methodology: Current average proficiency 3.5, '
        'required proficiency 4.0, gap score -0.5. Priority: Medium.'
    )
    doc.add_paragraph(
        'UX/UI Design: Current average proficiency 3.2, '
        'required proficiency 3.8, gap score -0.6. Priority: Medium.'
    )
    doc.add_paragraph(
        'Project Management (PMP/PRINCE2): Current average proficiency 3.4, '
        'required proficiency 4.0, gap score -0.6. Priority: Medium.'
    )
    doc.add_paragraph(
        'Executive Leadership: Current average proficiency 3.1, '
        'required proficiency 4.5, gap score -1.4. Priority: High.'
    )
    doc.add_paragraph(
        'Technical Writing & Documentation: Current average proficiency 3.6, '
        'required proficiency 3.8, gap score -0.2. Priority: Low.'
    )

    # ===== SECTION 5: HIRING PLAN =====
    doc.add_heading('5. Three-Year Hiring Plan (2026-2028)', level=1)
    doc.add_paragraph(
        'The hiring plan addresses both replacement needs (retirement and attrition) and growth '
        'positions aligned with strategic objectives. Quarterly targets account for seasonal '
        'recruitment cycles and business planning milestones.'
    )

    doc.add_heading('2026 Quarterly Hiring Targets', level=2)
    doc.add_paragraph(
        'Engineering: Q1 8, Q2 12, Q3 10, Q4 6, Annual Total 36. '
        'Product Management: Q1 2, Q2 4, Q3 3, Q4 2, Annual Total 11. '
        'Sales & Marketing: Q1 5, Q2 8, Q3 6, Q4 4, Annual Total 23. '
        'Research & Development: Q1 4, Q2 6, Q3 5, Q4 3, Annual Total 18. '
        'Operations: Q1 6, Q2 8, Q3 7, Q4 5, Annual Total 26. '
        'IT Infrastructure: Q1 3, Q2 4, Q3 3, Q4 2, Annual Total 12. '
        'Customer Success: Q1 2, Q2 3, Q3 3, Q4 2, Annual Total 10.'
    )

    doc.add_heading('2027 Quarterly Hiring Targets', level=2)
    doc.add_paragraph(
        'Engineering: Q1 10, Q2 14, Q3 12, Q4 8, Annual Total 44. '
        'Product Management: Q1 3, Q2 5, Q3 4, Q4 2, Annual Total 14. '
        'Sales & Marketing: Q1 6, Q2 10, Q3 8, Q4 5, Annual Total 29. '
        'Research & Development: Q1 5, Q2 8, Q3 6, Q4 4, Annual Total 23. '
        'Operations: Q1 4, Q2 6, Q3 5, Q4 3, Annual Total 18. '
        'IT Infrastructure: Q1 4, Q2 5, Q3 4, Q4 3, Annual Total 16. '
        'Customer Success: Q1 3, Q2 4, Q3 3, Q4 2, Annual Total 12.'
    )

    doc.add_heading('2028 Quarterly Hiring Targets', level=2)
    doc.add_paragraph(
        'Engineering: Q1 12, Q2 15, Q3 13, Q4 10, Annual Total 50. '
        'Product Management: Q1 3, Q2 5, Q3 4, Q4 3, Annual Total 15. '
        'Sales & Marketing: Q1 7, Q2 11, Q3 9, Q4 6, Annual Total 33. '
        'Research & Development: Q1 6, Q2 9, Q3 7, Q4 5, Annual Total 27. '
        'Operations: Q1 3, Q2 5, Q3 4, Q4 3, Annual Total 15. '
        'IT Infrastructure: Q1 4, Q2 6, Q3 5, Q4 3, Annual Total 18. '
        'Customer Success: Q1 3, Q2 5, Q3 4, Q4 3, Annual Total 15.'
    )

    # ===== SECTION 6: CONCLUSION =====
    doc.add_heading('6. Conclusion and Next Steps', level=1)
    doc.add_paragraph(
        'The data presented in this plan highlights both the urgency and the opportunity before '
        'Meridian Technologies. To fully leverage this analysis, the workforce planning team recommends '
        'converting the raw data above into structured tables for executive presentation and ongoing '
        'tracking. Each data section should be formatted with clear column headers, appropriate '
        'subtotals, and figure captions for cross-referencing. A list of tables at the document\'s '
        'beginning would provide quick navigation to each analytical view.'
    )
    doc.add_paragraph(
        'Immediate next steps include: (1) formatting all data sections into professional tables, '
        '(2) adding figure captions to each table for reference in board presentations, '
        '(3) creating a consolidated list of tables, and (4) distributing the finalized document '
        'to the executive leadership team for review at the Q1 2026 strategy meeting.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
