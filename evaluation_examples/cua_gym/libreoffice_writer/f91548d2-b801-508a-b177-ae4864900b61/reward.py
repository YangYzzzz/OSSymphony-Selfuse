"""
Reward Script: Add two new rows below the last row of the meeting agenda table
Task ID: writer_tbl_004
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Table has exactly 6 rows (4 original + 2 new)
  Component 2 (0.3): Row 5 (index 4) = 'Budget Review' | 'CFO' | '3:00 PM'
  Component 3 (0.3): Row 6 (index 5) = 'Closing Remarks' | 'CEO' | '3:30 PM'
  Total: 1.0
"""

import os

from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_004'

# Expected new row contents (from task_config.json ground truth)
EXPECTED_ROW5 = ['Budget Review', 'CFO', '3:00 PM']
EXPECTED_ROW6 = ['Closing Remarks', 'CEO', '3:30 PM']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document — if this fails, we cannot verify anything
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: document must contain at least one table
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    actual_row_count = len(table.rows)
    actual_col_count = len(table.columns)

    # Component 1: Table has exactly 6 rows (was 4; 2 new rows added) — 0.4 points
    # This FAILS on initial (4 rows) and PASSES on golden (6 rows).
    try:
        if actual_row_count == 6 and actual_col_count == 3:
            print(f"PASS: Component 1 — Table has {actual_row_count} rows x {actual_col_count} cols (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Expected 6 rows x 3 cols, found {actual_row_count} rows x {actual_col_count} cols")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Row 5 (index 4) = 'Budget Review' | 'CFO' | '3:00 PM' — 0.3 points
    # This FAILS on initial (row index 4 doesn't exist) and PASSES on golden.
    try:
        if actual_row_count >= 5:
            row5_cells = [cell.text.strip() for cell in table.rows[4].cells]
            if row5_cells == EXPECTED_ROW5:
                print(f"PASS: Component 2 — Row 5 = {row5_cells} (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Row 5 expected {EXPECTED_ROW5}, found {row5_cells}")
        else:
            print(f"FAIL: Component 2 — Table only has {actual_row_count} rows; row index 4 does not exist")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Row 6 (index 5) = 'Closing Remarks' | 'CEO' | '3:30 PM' — 0.3 points
    # This FAILS on initial (row index 5 doesn't exist) and PASSES on golden.
    try:
        if actual_row_count >= 6:
            row6_cells = [cell.text.strip() for cell in table.rows[5].cells]
            if row6_cells == EXPECTED_ROW6:
                print(f"PASS: Component 3 — Row 6 = {row6_cells} (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 3 — Row 6 expected {EXPECTED_ROW6}, found {row6_cells}")
        else:
            print(f"FAIL: Component 3 — Table only has {actual_row_count} rows; row index 5 does not exist")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
