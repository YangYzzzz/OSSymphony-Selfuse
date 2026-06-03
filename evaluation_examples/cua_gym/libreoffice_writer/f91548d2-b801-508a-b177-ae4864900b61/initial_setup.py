"""
Initial Setup: Meeting agenda document with a 4-row table
Task ID: writer_tbl_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_tbl_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DESKTOP_OUTPUT = f'{WORKDIR}/Desktop/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a heading for the document
    heading = doc.add_heading('Meeting Agenda', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a short intro paragraph
    doc.add_paragraph('The following agenda outlines the schedule for the quarterly review meeting.')

    # Create the meeting agenda table with 4 rows and 3 columns
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Row 1: Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Topic'
    header_cells[1].text = 'Presenter'
    header_cells[2].text = 'Time'

    # Make header row bold
    for cell in header_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Row 2
    row2 = table.rows[1].cells
    row2[0].text = 'Welcome'
    row2[1].text = 'CEO'
    row2[2].text = '1:00 PM'

    # Row 3
    row3 = table.rows[2].cells
    row3[0].text = 'Q3 Results'
    row3[1].text = 'VP Sales'
    row3[2].text = '1:30 PM'

    # Row 4
    row4 = table.rows[3].cells
    row4[0].text = 'Marketing Plan'
    row4[1].text = 'CMO'
    row4[2].text = '2:15 PM'

    # Add a footer note
    doc.add_paragraph('')
    doc.add_paragraph('Note: All sessions are held in Conference Room B.')

    # Ensure Desktop directory exists
    desktop_dir = f'{WORKDIR}/Desktop'
    os.makedirs(desktop_dir, exist_ok=True)

    # Save to both locations
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Also save to Desktop as mentioned in task context
    doc.save(DESKTOP_OUTPUT)
    print(f'Initial file also saved to Desktop: {DESKTOP_OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
