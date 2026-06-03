"""
Initial Setup: Configure footnote numbering restart per chapter
Task ID: writer_acad_024
Domain: libreoffice_writer

Creates a thesis document with 3 chapters, each on a new page,
with continuous footnote numbering (1-15).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_024'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_footnote(paragraph, footnote_id, footnote_text, doc):
    """Add a footnote to a paragraph using XML manipulation."""
    # Get or create footnotes part
    footnotes_part = None
    for rel in doc.part.rels.values():
        if 'footnotes' in rel.reltype:
            footnotes_part = rel.target_part
            break

    if footnotes_part is None:
        # Create footnotes part from scratch
        from docx.opc.part import Part
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        import io

        footnotes_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            '             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '  <w:footnote w:type="separator" w:id="-1">'
            '    <w:p><w:r><w:separator/></w:r></w:p>'
            '  </w:footnote>'
            '  <w:footnote w:type="continuationSeparator" w:id="0">'
            '    <w:p><w:r><w:continuationSeparator/></w:r></w:p>'
            '  </w:footnote>'
            '</w:footnotes>'
        )
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        footnotes_part = Part(
            PackURI('/word/footnotes.xml'),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml',
            footnotes_xml.encode('utf-8'),
            doc.part.package,
        )
        doc.part.relate_to(
            footnotes_part,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes',
        )

    # Parse existing footnotes XML
    footnotes_elem = etree.fromstring(footnotes_part.blob)

    # Create new footnote element
    fn = etree.SubElement(footnotes_elem, qn('w:footnote'))
    fn.set(qn('w:id'), str(footnote_id))
    fn_p = etree.SubElement(fn, qn('w:p'))

    # Footnote reference run (the superscript number inside the footnote)
    fn_r_ref = etree.SubElement(fn_p, qn('w:r'))
    fn_rpr_ref = etree.SubElement(fn_r_ref, qn('w:rPr'))
    etree.SubElement(fn_rpr_ref, qn('w:rStyle')).set(qn('w:val'), 'FootnoteReference')
    etree.SubElement(fn_r_ref, qn('w:footnoteRef'))

    # Space after number
    fn_r_space = etree.SubElement(fn_p, qn('w:r'))
    fn_t_space = etree.SubElement(fn_r_space, qn('w:t'))
    fn_t_space.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    fn_t_space.text = ' '

    # Footnote text run
    fn_r_text = etree.SubElement(fn_p, qn('w:r'))
    fn_t = etree.SubElement(fn_r_text, qn('w:t'))
    fn_t.text = footnote_text

    # Update the part blob
    footnotes_part._blob = etree.tostring(footnotes_elem, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Add footnote reference in the body paragraph
    run = paragraph.add_run()
    rpr = run._element.get_or_add_rPr()
    etree.SubElement(rpr, qn('w:rStyle')).set(qn('w:val'), 'FootnoteReference')
    fn_ref = etree.SubElement(run._element, qn('w:footnoteRef'))
    # Actually we need a footnoteReference, not footnoteRef in body
    run._element.remove(fn_ref)
    fn_reference = etree.SubElement(run._element, qn('w:footnoteReference'))
    fn_reference.set(qn('w:id'), str(footnote_id))


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Configure heading styles
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Ensure FootnoteReference style exists
    # (python-docx may not have it, we'll rely on the XML)

    # =============================================
    # CHAPTER 1: Introduction to Machine Learning
    # =============================================
    h1 = doc.add_heading('Chapter 1: Introduction to Machine Learning', level=1)

    p1 = doc.add_paragraph(
        'Machine learning is a subset of artificial intelligence that focuses on '
        'building systems capable of learning from data and improving their performance '
        'over time without being explicitly programmed.'
    )
    add_footnote(p1, 1, 'Russell, S. & Norvig, P. (2021). Artificial Intelligence: A Modern Approach, 4th ed. Pearson.', doc)

    p2 = doc.add_paragraph(
        'The field has experienced remarkable growth since the early 2010s, driven largely '
        'by advances in computational power and the availability of large-scale datasets.'
    )
    add_footnote(p2, 2, 'LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep learning. Nature, 521(7553), 436-444.', doc)

    p3 = doc.add_paragraph(
        'Supervised learning remains the most widely used paradigm, where models learn '
        'from labeled training examples to make predictions on unseen data. Common algorithms '
        'include linear regression, decision trees, and neural networks.'
    )
    add_footnote(p3, 3, 'Bishop, C. M. (2006). Pattern Recognition and Machine Learning. Springer.', doc)

    p4 = doc.add_paragraph(
        'The bias-variance tradeoff is a fundamental concept that describes the tension '
        'between a model\'s ability to fit training data and its capacity to generalize '
        'to new observations.'
    )
    add_footnote(p4, 4, 'Hastie, T., Tibshirani, R., & Friedman, J. (2009). The Elements of Statistical Learning. Springer.', doc)

    p5 = doc.add_paragraph(
        'Cross-validation techniques, particularly k-fold cross-validation, provide robust '
        'methods for estimating model performance and selecting hyperparameters.'
    )
    add_footnote(p5, 5, 'Kohavi, R. (1995). A study of cross-validation and bootstrap for accuracy estimation. IJCAI, 14(2), 1137-1145.', doc)

    # =============================================
    # CHAPTER 2: Deep Learning Architectures
    # =============================================
    # Add section break for new page
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    h2 = doc.add_heading('Chapter 2: Deep Learning Architectures', level=1)

    p6 = doc.add_paragraph(
        'Convolutional neural networks (CNNs) have revolutionized computer vision tasks, '
        'achieving superhuman performance in image classification benchmarks such as ImageNet.'
    )
    add_footnote(p6, 6, 'Krizhevsky, A., Sutskever, I., & Hinton, G. E. (2012). ImageNet classification with deep CNNs. NeurIPS, 25.', doc)

    p7 = doc.add_paragraph(
        'Recurrent neural networks, particularly Long Short-Term Memory (LSTM) networks, '
        'were the dominant architecture for sequential data processing before the advent '
        'of transformers.'
    )
    add_footnote(p7, 7, 'Hochreiter, S. & Schmidhuber, J. (1997). Long short-term memory. Neural Computation, 9(8), 1735-1780.', doc)

    p8 = doc.add_paragraph(
        'The transformer architecture, introduced in 2017, replaced recurrence with '
        'self-attention mechanisms and has since become the foundation for state-of-the-art '
        'natural language processing models.'
    )
    add_footnote(p8, 8, 'Vaswani, A. et al. (2017). Attention is all you need. NeurIPS, 30.', doc)

    p9 = doc.add_paragraph(
        'Generative adversarial networks (GANs) consist of a generator and discriminator '
        'trained in an adversarial fashion, producing remarkably realistic synthetic data '
        'including images, audio, and text.'
    )
    add_footnote(p9, 9, 'Goodfellow, I. et al. (2014). Generative adversarial nets. NeurIPS, 27.', doc)

    p10 = doc.add_paragraph(
        'Transfer learning has emerged as a practical strategy where models pre-trained '
        'on large datasets are fine-tuned for specific downstream tasks, significantly '
        'reducing the data and compute requirements.'
    )
    add_footnote(p10, 10, 'Devlin, J. et al. (2019). BERT: Pre-training of deep bidirectional transformers. NAACL-HLT, 4171-4186.', doc)

    # =============================================
    # CHAPTER 3: Ethics and Future Directions
    # =============================================
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    h3 = doc.add_heading('Chapter 3: Ethics and Future Directions', level=1)

    p11 = doc.add_paragraph(
        'Algorithmic bias represents one of the most pressing ethical challenges in machine '
        'learning, as models trained on historical data can perpetuate and amplify existing '
        'societal inequalities.'
    )
    add_footnote(p11, 11, 'Barocas, S. & Selbst, A. D. (2016). Big data\'s disparate impact. California Law Review, 104, 671-732.', doc)

    p12 = doc.add_paragraph(
        'Explainable AI (XAI) seeks to make machine learning models more interpretable '
        'and transparent, enabling stakeholders to understand how decisions are made.'
    )
    add_footnote(p12, 12, 'Ribeiro, M. T. et al. (2016). "Why should I trust you?": Explaining the predictions of any classifier. KDD, 1135-1144.', doc)

    p13 = doc.add_paragraph(
        'The environmental impact of training large-scale models has become a growing '
        'concern, with some estimates suggesting that training a single large language '
        'model can emit as much carbon as five cars over their lifetimes.'
    )
    add_footnote(p13, 13, 'Strubell, E., Ganesh, A., & McCallum, A. (2019). Energy and policy considerations for deep learning in NLP. ACL, 3645-3650.', doc)

    p14 = doc.add_paragraph(
        'Federated learning offers a privacy-preserving approach where models are trained '
        'across decentralized devices without sharing raw data, addressing data sovereignty '
        'and privacy regulations.'
    )
    add_footnote(p14, 14, 'McMahan, B. et al. (2017). Communication-efficient learning of deep networks from decentralized data. AISTATS, 1273-1282.', doc)

    p15 = doc.add_paragraph(
        'The development of artificial general intelligence (AGI) remains an open research '
        'question, with significant debate about timelines, feasibility, and the potential '
        'risks associated with superintelligent systems.'
    )
    add_footnote(p15, 15, 'Bostrom, N. (2014). Superintelligence: Paths, Dangers, Strategies. Oxford University Press.', doc)

    # Footnotes are continuously numbered 1-15 (default behavior)
    # Do NOT set restart per section -- that's the task for the agent

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
