"""
Reward Script: Create Avery 5066 file folder labels with department names
Task ID: writer_lec_055
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table exists with label grid structure
  Component 2 (0.35): Six department names present in correct cells
  Component 3 (0.25): Department labels use bold 12pt font
  Component 4 (0.15): Remaining label cells are empty
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_055'

# Expected department names
EXPECTED_DEPARTMENTS = [
    'Human Resources', 'Accounting', 'Marketing',
    'IT Support', 'Legal', 'Operations'
]


def persist_app_state(domain: str):
    """Attempt to save any unsaved LibreOffice state."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        print(f"CRITICAL: Missing dependency: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with label grid structure (0.25 points)
    # A label document should have at least one table acting as the label grid.
    # Avery 5066 labels are arranged in a grid (typically 3 columns, multiple rows).
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            # Avery 5066: expect 3 columns and multiple rows (at least 2 rows to hold 6 labels)
            if num_cols == 3 and num_rows >= 2:
                print(f"PASS: Component 1 -- Table found with {num_rows} rows x {num_cols} cols (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 -- Table has {num_rows} rows x {num_cols} cols, expected 3 cols and >= 2 rows")
        else:
            print(f"FAIL: Component 1 -- No tables found in document (expected label grid)")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Six department names present in correct cells (0.35 points)
    # The six departments should be in the first 6 cells of the table (row-major order).
    try:
        table = doc.tables[0]
        found_departments = []
        all_cell_texts = []

        # Collect all cell text in row-major order
        for row in table.rows:
            for cell in row.cells:
                all_cell_texts.append(cell.text.strip())

        # Check that all 6 department names appear in the first 6 cells
        matches = 0
        for i, dept in enumerate(EXPECTED_DEPARTMENTS):
            if i < len(all_cell_texts) and all_cell_texts[i].lower() == dept.lower():
                matches += 1
                found_departments.append(dept)

        if matches == 6:
            print(f"PASS: Component 2 -- All 6 departments found in correct order (0.35 pts)")
            total_score += 0.35
        elif matches >= 4:
            # Partial credit: at least 4 of 6 departments in correct position
            partial = 0.35 * (matches / 6)
            print(f"PARTIAL: Component 2 -- {matches}/6 departments in correct positions ({partial:.2f} pts)")
            total_score += partial
        elif matches >= 1:
            # Minimal credit for some matches
            partial = 0.35 * (matches / 6)
            print(f"PARTIAL: Component 2 -- {matches}/6 departments found ({partial:.2f} pts)")
            total_score += partial
        else:
            # Check if departments exist anywhere in the table
            all_text_joined = ' '.join(all_cell_texts).lower()
            found_any = sum(1 for d in EXPECTED_DEPARTMENTS if d.lower() in all_text_joined)
            if found_any >= 4:
                partial = 0.35 * 0.5
                print(f"PARTIAL: Component 2 -- {found_any} departments found but not in expected positions ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 -- Expected departments {EXPECTED_DEPARTMENTS}, found cells: {all_cell_texts[:8]}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Department labels use bold 12pt font (0.25 points)
    # Each department label should be bold and 12pt (152400 EMU).
    try:
        table = doc.tables[0]
        dept_cells_checked = 0
        bold_and_sized = 0

        cell_idx = 0
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text.lower() in [d.lower() for d in EXPECTED_DEPARTMENTS]:
                    dept_cells_checked += 1
                    runs = []
                    for para in cell.paragraphs:
                        runs.extend(para.runs)

                    if not runs:
                        continue

                    # Check all runs in this cell for bold and size
                    all_bold = all(r.font.bold is True for r in runs if r.text.strip())
                    all_12pt = all(
                        r.font.size is not None and abs(r.font.size.pt - 12.0) < 0.5
                        for r in runs if r.text.strip()
                    )

                    if all_bold and all_12pt:
                        bold_and_sized += 1

                cell_idx += 1

        if dept_cells_checked >= 6 and bold_and_sized >= 6:
            print(f"PASS: Component 3 -- All 6 department labels are bold 12pt (0.25 pts)")
            total_score += 0.25
        elif bold_and_sized >= 4:
            partial = 0.25 * (bold_and_sized / 6)
            print(f"PARTIAL: Component 3 -- {bold_and_sized}/6 labels are bold 12pt ({partial:.2f} pts)")
            total_score += partial
        elif dept_cells_checked == 0:
            print(f"FAIL: Component 3 -- No department labels found to check formatting")
        else:
            print(f"FAIL: Component 3 -- Only {bold_and_sized}/{dept_cells_checked} labels have bold 12pt font")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Remaining label cells are empty (0.15 points)
    # After the 6 department labels, remaining cells in the table should be blank.
    try:
        table = doc.tables[0]
        non_dept_cells = []
        dept_set = set(d.lower() for d in EXPECTED_DEPARTMENTS)

        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text.lower() not in dept_set:
                    non_dept_cells.append(text)

        # All non-department cells should be empty
        non_empty_extra = [t for t in non_dept_cells if t]
        total_non_dept = len(non_dept_cells)

        if total_non_dept > 0 and len(non_empty_extra) == 0:
            print(f"PASS: Component 4 -- {total_non_dept} remaining cells are all empty (0.15 pts)")
            total_score += 0.15
        elif total_non_dept == 0:
            # Edge case: every cell is a department (unlikely for label sheet)
            print(f"FAIL: Component 4 -- No remaining cells found (all cells contain departments)")
        else:
            print(f"FAIL: Component 4 -- {len(non_empty_extra)} remaining cells are non-empty: {non_empty_extra[:5]}")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
