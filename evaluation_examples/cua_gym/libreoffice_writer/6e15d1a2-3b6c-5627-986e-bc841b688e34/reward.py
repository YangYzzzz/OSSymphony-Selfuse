"""
Reward Script: Weekly Status Report Template in LibreOffice Writer
Task ID: writer_wf_022
Domain: libreoffice_writer
Scoring:
  Component 1: Landscape orientation (0.10)
  Component 2: Title 'Weekly Status Report' with Title style (0.15)
  Component 3: Info fields present (Project Name, Report Period, Prepared By, Date) (0.10)
  Component 4: Four Heading 2 sections (Accomplishments, Planned, Issues, Key Metrics) (0.20)
  Component 5: Bulleted lists under Accomplishments and Planned sections (0.10)
  Component 6: Issues and Risks table with correct structure (header + 3 rows, 3 cols) (0.20)
  Component 7: Key Metrics table with correct structure (header + 4 rows, 4 cols) (0.15)
"""

import os
from docx import Document
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_022'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Landscape orientation (0.10 points)
    try:
        section = doc.sections[0]
        is_landscape = section.orientation == WD_ORIENT.LANDSCAPE
        # Also check width > height as a secondary confirmation
        width_gt_height = section.page_width > section.page_height
        if is_landscape and width_gt_height:
            print(f"PASS: Component 1 — Landscape orientation confirmed (0.10 pts)")
            total_score += 0.10
        elif is_landscape or width_gt_height:
            print(f"PARTIAL: Component 1 — Partial landscape (orient={section.orientation}, w={section.page_width}, h={section.page_height}) (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 1 — Portrait orientation (orient={section.orientation}, w={section.page_width}, h={section.page_height})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Collect paragraph info for subsequent checks
    paragraphs = doc.paragraphs
    para_styles = [(p.style.name if p.style else 'None', p.text.strip()) for p in paragraphs]

    # Component 2: Title 'Weekly Status Report' with Title style (0.15 points)
    try:
        title_found = False
        for style_name, text in para_styles:
            if 'weekly status report' in text.lower():
                if style_name == 'Title':
                    print(f"PASS: Component 2 — Title 'Weekly Status Report' with Title style (0.15 pts)")
                    total_score += 0.15
                    title_found = True
                elif 'heading' in style_name.lower():
                    print(f"PARTIAL: Component 2 — Title text found but style is '{style_name}' not 'Title' (0.08 pts)")
                    total_score += 0.08
                    title_found = True
                else:
                    print(f"PARTIAL: Component 2 — Title text found but style is '{style_name}' (0.05 pts)")
                    total_score += 0.05
                    title_found = True
                break
        if not title_found:
            print(f"FAIL: Component 2 — No paragraph containing 'Weekly Status Report' found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Info fields (Project Name, Report Period, Prepared By, Date) (0.10 points)
    try:
        all_text_lower = ' '.join(text.lower() for _, text in para_styles)
        required_fields = ['project name', 'report period', 'prepared by', 'date']
        fields_found = sum(1 for f in required_fields if f in all_text_lower)
        if fields_found == 4:
            print(f"PASS: Component 3 — All 4 info fields found (0.10 pts)")
            total_score += 0.10
        elif fields_found >= 2:
            pts = round(0.10 * fields_found / 4, 2)
            print(f"PARTIAL: Component 3 — {fields_found}/4 info fields found ('{', '.join(f for f in required_fields if f in all_text_lower)}') ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 3 — Only {fields_found}/4 info fields found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Four Heading 2 sections (0.20 points)
    try:
        expected_headings = ['accomplishments this week', 'planned for next week', 'issues and risks', 'key metrics']
        heading2_texts = [text.lower() for style_name, text in para_styles if style_name == 'Heading 2']
        headings_found = 0
        for eh in expected_headings:
            if any(eh in h for h in heading2_texts):
                headings_found += 1
        if headings_found == 4:
            print(f"PASS: Component 4 — All 4 Heading 2 sections found (0.20 pts)")
            total_score += 0.20
        elif headings_found >= 1:
            pts = round(0.20 * headings_found / 4, 2)
            print(f"PARTIAL: Component 4 — {headings_found}/4 Heading 2 sections found ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 4 — No expected Heading 2 sections found (found: {heading2_texts})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Bulleted lists under Accomplishments and Planned sections (0.10 points)
    try:
        # Find bullet items after each heading
        bullet_styles = {'List Bullet', 'List Bullet 2', 'List Bullet 3'}
        accomplishments_bullets = 0
        planned_bullets = 0
        current_section = None
        for style_name, text in para_styles:
            if style_name == 'Heading 2':
                if 'accomplishments' in text.lower():
                    current_section = 'accomplishments'
                elif 'planned' in text.lower():
                    current_section = 'planned'
                else:
                    current_section = 'other'
            elif current_section == 'accomplishments' and style_name in bullet_styles:
                accomplishments_bullets += 1
            elif current_section == 'planned' and style_name in bullet_styles:
                planned_bullets += 1

        has_acc_bullets = accomplishments_bullets >= 1
        has_plan_bullets = planned_bullets >= 1
        if has_acc_bullets and has_plan_bullets:
            print(f"PASS: Component 5 — Bullets found: Accomplishments={accomplishments_bullets}, Planned={planned_bullets} (0.10 pts)")
            total_score += 0.10
        elif has_acc_bullets or has_plan_bullets:
            print(f"PARTIAL: Component 5 — Only one section has bullets (Acc={accomplishments_bullets}, Plan={planned_bullets}) (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 — No bullet lists found (Acc={accomplishments_bullets}, Plan={planned_bullets})")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Issues and Risks table (header + 3 data rows, 3 cols: Issue, Impact, Mitigation) (0.20 points)
    try:
        tables = doc.tables
        issues_table_found = False
        if len(tables) >= 1:
            # The first table should be Issues and Risks
            for t in tables:
                header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
                if 'issue' in header_cells and 'impact' in header_cells and 'mitigation' in header_cells:
                    issues_table_found = True
                    num_rows = len(t.rows)
                    num_cols = len(t.columns)
                    pts = 0.0
                    # Check columns: should have 3 (Issue, Impact, Mitigation)
                    if num_cols == 3:
                        pts += 0.08
                        print(f"  Issues table: 3 columns confirmed")
                    else:
                        print(f"  Issues table: expected 3 cols, found {num_cols}")
                    # Check rows: header + 3 data rows = 4 total
                    if num_rows >= 4:
                        pts += 0.08
                        print(f"  Issues table: {num_rows} rows (>= 4 expected)")
                    elif num_rows >= 2:
                        pts += 0.04
                        print(f"  Issues table: only {num_rows} rows (expected >= 4)")
                    # Check data rows have content
                    data_rows_with_content = sum(1 for r in t.rows[1:] if any(c.text.strip() for c in r.cells))
                    if data_rows_with_content >= 3:
                        pts += 0.04
                        print(f"  Issues table: {data_rows_with_content} data rows with content")
                    if pts > 0:
                        print(f"PASS: Component 6 — Issues and Risks table found ({pts} pts)")
                        total_score += pts
                    else:
                        print(f"FAIL: Component 6 — Issues table structure incorrect")
                    break
        if not issues_table_found:
            print(f"FAIL: Component 6 — No Issues and Risks table found (tables={len(tables)})")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Key Metrics table (header + 4 data rows, 4 cols: Metric, Target, Actual, Status) (0.15 points)
    try:
        tables = doc.tables
        metrics_table_found = False
        for t in tables:
            header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
            if 'metric' in header_cells and 'target' in header_cells and 'actual' in header_cells and 'status' in header_cells:
                metrics_table_found = True
                num_rows = len(t.rows)
                num_cols = len(t.columns)
                pts = 0.0
                # Check columns: should have 4
                if num_cols == 4:
                    pts += 0.05
                    print(f"  Metrics table: 4 columns confirmed")
                else:
                    print(f"  Metrics table: expected 4 cols, found {num_cols}")
                # Check rows: header + 4 data rows = 5 total
                if num_rows >= 5:
                    pts += 0.06
                    print(f"  Metrics table: {num_rows} rows (>= 5 expected)")
                elif num_rows >= 2:
                    pts += 0.03
                    print(f"  Metrics table: only {num_rows} rows (expected >= 5)")
                # Check data rows have content
                data_rows_with_content = sum(1 for r in t.rows[1:] if any(c.text.strip() for c in r.cells))
                if data_rows_with_content >= 4:
                    pts += 0.04
                    print(f"  Metrics table: {data_rows_with_content} data rows with content")
                if pts > 0:
                    print(f"PASS: Component 7 — Key Metrics table found ({pts} pts)")
                    total_score += pts
                else:
                    print(f"FAIL: Component 7 — Metrics table structure incorrect")
                break
        if not metrics_table_found:
            print(f"FAIL: Component 7 — No Key Metrics table found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()

import time
time.sleep(0.5)

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
