"""
Initial Setup: Create a membership notice document with member status text
Task ID: writer_mt_027
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_027'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Greenfield Community Association', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_heading('Membership Status Notice', level=1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Date line ---
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = date_para.add_run('Date: April 15, 2025')
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # --- Greeting ---
    greeting = doc.add_paragraph()
    run = greeting.add_run('Dear Member,')
    run.font.size = Pt(12)
    run.bold = True

    # --- Body paragraphs ---
    body1 = doc.add_paragraph()
    run = body1.add_run(
        'Thank you for being a valued member of the Greenfield Community Association. '
        'This notice is to inform you about your current membership status. '
        'Please review the information below carefully and take any necessary action.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    body2 = doc.add_paragraph()
    run = body2.add_run(
        'Based on our records, the following details are associated with your membership:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Member details table ---
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    fields = [
        ('Field', 'Value'),
        ('Member Name', '{MemberName}'),
        ('Member ID', '{MemberID}'),
        ('Membership Type', '{MembershipType}'),
        ('Membership Expiry', '{MembershipExpiry}'),
    ]
    for i, (label, value) in enumerate(fields):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        # Header row formatting
        if i == 0:
            run_l = cell_label.paragraphs[0].add_run(label)
            run_l.bold = True
            run_l.font.size = Pt(11)
            run_v = cell_value.paragraphs[0].add_run(value)
            run_v.bold = True
            run_v.font.size = Pt(11)
        else:
            run_l = cell_label.paragraphs[0].add_run(label)
            run_l.font.size = Pt(11)
            run_v = cell_value.paragraphs[0].add_run(value)
            run_v.font.size = Pt(11)
            run_v.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

    doc.add_paragraph()  # spacer

    # --- Status section ---
    status_heading = doc.add_heading('Membership Status:', level=2)

    status_para = doc.add_paragraph()
    run = status_para.add_run(
        'Your membership status will be determined based on your expiry date. '
        'Please insert the appropriate conditional merge field below to display '
        'the correct status message.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # Placeholder line where the conditional field should go
    placeholder = doc.add_paragraph()
    run = placeholder.add_run('Status: ')
    run.font.size = Pt(12)
    run.bold = True
    run2 = placeholder.add_run('[Insert conditional field here]')
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run2.italic = True

    doc.add_paragraph()  # spacer

    # --- Action items ---
    action_heading = doc.add_heading('Required Actions:', level=2)

    actions = [
        'If your membership shows "Renewal Required", please visit the front desk '
        'or our website at www.greenfieldca.org/renew to complete your renewal.',
        'Members with "Active Member" status do not need to take any action at this time.',
        'All renewals must be completed within 30 days of receiving this notice to '
        'avoid any interruption in membership benefits.',
        'For questions about your membership status, contact us at '
        'membership@greenfieldca.org or call (555) 234-8901.',
    ]
    for action in actions:
        para = doc.add_paragraph(action, style='List Bullet')
        for run in para.runs:
            run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # --- Closing ---
    closing = doc.add_paragraph()
    run = closing.add_run(
        'We appreciate your continued support and look forward to serving you.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    sign = doc.add_paragraph()
    run = sign.add_run('Sincerely,')
    run.font.size = Pt(11)

    name_para = doc.add_paragraph()
    run = name_para.add_run('Patricia Hawkins')
    run.font.size = Pt(11)
    run.bold = True

    title_para = doc.add_paragraph()
    run = title_para.add_run('Membership Director')
    run.font.size = Pt(11)

    org_para = doc.add_paragraph()
    run = org_para.add_run('Greenfield Community Association')
    run.font.size = Pt(11)

    # --- Footer info ---
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fp.add_run('Greenfield Community Association | 450 Oak Street, Greenfield, CA 93927 | (555) 234-8901')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
