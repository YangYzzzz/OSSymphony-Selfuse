"""
FINAL REWARD SCRIPT - SUCCESS
Task: I just pasted a huge chunk of text into LibreOffice Writer 7.6 and it left me with pages full of “empty” lines that actually contain nothing but space characters. When I switch on non-printing marks, each of these paragraphs shows up as a bunch of small blue dots followed by a ¶. I need to get rid of every paragraph that matches the regex pattern "^[ ]+$" (one or more spaces, no tabs or other characters) in one sweep—basically a single Replace All run instead of deleting them line by line. How do I pull that off?
Generated: 2025-09-10 16:40:59
Status: success
Model: azure-o3
Total Steps: 5
"""

import os
import re
from docx import Document

def verify_remove_spaces_only_paragraphs(file_path: str) -> float:
    """Reward script for LibreOffice Writer cleaning task.

    The task: remove every paragraph that contains ONLY one or more space
    characters (regex: ^[ ]+$) in one sweep. This script verifies the result
    by inspecting the DOCX produced after the agent's actions.

    Scoring strategy (progressive, 0.0‒1.0):
      • 1.0  – NO paragraphs consisting solely of spaces remain.
      • <1.0 – Proportional to the fraction of paragraphs that are *not*
                space-only (cleaned_ratio). Example: if half of such
                paragraphs remain → score ≈ 0.5.
      • 0.0  – File missing / cannot load / zero paragraphs.
    """

    print(f"Verifying cleaned document: {file_path}")

    # Prerequisite checks (NO points awarded here)
    if not os.path.exists(file_path):
        print(f"✗ File not found: {file_path}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
        print(f"✓ Document loaded successfully with {len(doc.paragraphs)} paragraphs")
    except Exception as e:
        print(f"✗ Failed to load DOCX: {e}")
        print("REWARD: 0.0")
        return 0.0

    total_paragraphs = len(doc.paragraphs)
    spaces_only_count = 0

    # REAL verification: inspect every paragraph's text
    for p in doc.paragraphs:
        if re.fullmatch(r"[ ]+", p.text):
            spaces_only_count += 1

    print(f"Total paragraphs: {total_paragraphs}")
    print(f"Paragraphs containing ONLY space characters: {spaces_only_count}")

    # Progressive scoring based on how thoroughly the cleanup was done
    if total_paragraphs == 0:
        print("✗ Document contains no paragraphs – cannot verify cleaning task")
        final_score = 0.0
    else:
        cleaned_ratio = (total_paragraphs - spaces_only_count) / total_paragraphs
        if spaces_only_count == 0:
            print("✓ No space-only paragraphs remain – perfect clean up!")
            final_score = 1.0
        else:
            final_score = round(cleaned_ratio, 2)
            print(f"Partial clean-up detected – awarding proportional score: {final_score}")

    print(f"REWARD: {final_score}")
    return final_score

# ------------------ MAIN EXECUTION ------------------
if __name__ == "__main__":
    TARGET_DOC = "/home/user/i_just_pasted_a_huge_chunk_of_text_into_libreoffice_writer_76_and_it_left_me_with_pages_full_of_empt.docx"
    verify_remove_spaces_only_paragraphs(TARGET_DOC)
