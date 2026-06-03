"""
FINAL REWARD SCRIPT - SUCCESS
Task: The opening paragraph looks crammed. How can I force a line break immediately after the very first sentence—right after the period—without creating a whole new paragraph in LibreOffice Writer?
Generated: 2025-09-10 16:06:13
Status: success
Model: azure-o3
Total Steps: 3
"""

import os
import sys
from docx import Document


def has_soft_line_break(paragraph):
    """Detects a manual (soft) line break inside a Word paragraph.

    The check is done in two ways:
    1. python-docx exposes a manual line break as a literal "\n" in paragraph.text.
    2. Low-level XML inspection for <w:br/> elements (WordprocessingML line break tags).
    """
    # Text-level detection
    if "\n" in paragraph.text:
        return True

    # XML-level detection (fallback – more robust)
    p_xml = paragraph._p  # lxml element of the paragraph
    for elem in p_xml.iter():
        if elem.tag.endswith("}br"):
            return True
    return False


def verify_line_break_task(file_path):
    """Reward script for the LibreOffice Writer task.

    Requirements to score full points (1.0):
    1. The first paragraph contains a manual line break immediately after the first
       sentence (0.7 points).
    2. The second sentence is *not* placed in its own paragraph – the opening still
       counts as a single paragraph in the document (0.3 points).
    Progressive scoring is applied so partial completion earns partial credit.
    """
    print(f"Starting verification for: {file_path}")
    score = 0.0

    # ------------------------------------------------------------------
    # Safety check – the file must exist and load (no points for this)
    # ------------------------------------------------------------------
    if not os.path.exists(file_path):
        print(f"✗ File does not exist: {file_path}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
        print(f"✓ Loaded DOCX successfully. Total paragraphs: {len(doc.paragraphs)}")
    except Exception as e:
        print(f"✗ Failed to load DOCX: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ------------------------------------------------------------------
    # Requirement 1: Manual line break after first sentence (0.7)
    # ------------------------------------------------------------------
    if not doc.paragraphs:
        print("✗ Document contains no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    first_para = doc.paragraphs[0]
    if has_soft_line_break(first_para):
        print("✓ Detected manual line break within the first paragraph (0.7 points)")
        score += 0.7
    else:
        print("✗ No manual line break detected in the first paragraph")

    # ------------------------------------------------------------------
    # Requirement 2: Opening remains a single paragraph (0.3)
    # ------------------------------------------------------------------
    if len(doc.paragraphs) == 1:
        print("✓ Opening remains a single paragraph (0.3 points)")
        score += 0.3
    else:
        print(f"✗ Found {len(doc.paragraphs)} paragraphs; expected 1 for opening")

    # ------------------------------------------------------------------
    # Final scoring
    # ------------------------------------------------------------------
    final_score = round(min(score, 1.0), 2)
    print(f"Total score: {final_score}")
    print(f"REWARD: {final_score}")
    return final_score


if __name__ == "__main__":
    # Default path used in the evaluation environment
    default_path = "/home/user/the_opening_paragraph_looks_crammed_how_can_i_force_a_line_break_immediately_after_the_very_first_se.docx"
    verify_line_break_task(default_path)
