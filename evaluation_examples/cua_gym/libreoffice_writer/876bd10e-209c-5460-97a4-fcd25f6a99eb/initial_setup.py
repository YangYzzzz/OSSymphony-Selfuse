"""
Initial Setup: Network Connectivity Troubleshooting Guide
Task ID: osworld_writer_blank_line_insertion_006
Domain: libreoffice_writer

Creates a technical troubleshooting guide document with 8 numbered steps
listed consecutively WITHOUT any blank lines between them.
The agent task is to add a blank line after each step to improve readability.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_blank_line_insertion_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading("Network Connectivity Troubleshooting Guide", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction
    intro = doc.add_paragraph(
        "This guide provides step-by-step instructions for diagnosing and resolving "
        "common network connectivity issues on workstations and laptops. Follow each "
        "step in order before proceeding to the next."
    )
    intro.paragraph_format.space_after = Pt(6)

    # Section heading
    doc.add_heading("Troubleshooting Steps", level=1)

    # 8 numbered steps listed consecutively WITHOUT blank lines between them
    steps = [
        "1. Check physical connections — Verify that the Ethernet cable is securely plugged into both the network port on the device and the wall socket or switch. If using Wi-Fi, ensure the wireless adapter is enabled and not in airplane mode.",
        "2. Restart the network adapter — Open Device Manager, expand the Network Adapters section, right-click your adapter, and select Disable. Wait five seconds, then right-click again and select Enable to reinitialize the hardware.",
        "3. Release and renew the IP address — Open Command Prompt as Administrator and run 'ipconfig /release' to drop the current IP lease. Once completed, run 'ipconfig /renew' to request a new address from the DHCP server.",
        "4. Flush the DNS cache — In the same Command Prompt window, type 'ipconfig /flushdns' and press Enter. This clears outdated or corrupted DNS resolver cache entries that may prevent domain name resolution.",
        "5. Ping the default gateway — Run 'ipconfig' to find the Default Gateway address (e.g., 192.168.1.1), then run 'ping 192.168.1.1'. If the ping fails, the problem is likely between the device and the local router.",
        "6. Test external connectivity — Ping a reliable external server such as Google's DNS at 8.8.8.8 using 'ping 8.8.8.8 -t'. Successful replies confirm basic internet routing works; timeouts indicate an upstream issue.",
        "7. Check firewall and security software — Temporarily disable the Windows Firewall and any third-party antivirus or VPN software, then retest connectivity. Re-enable all security software immediately after the test regardless of outcome.",
        "8. Update or reinstall network drivers — Navigate to the manufacturer's website, download the latest driver for your network adapter model, and run the installer. If issues persist, use Device Manager to uninstall the device and restart to trigger automatic driver reinstallation.",
    ]

    for step in steps:
        para = doc.add_paragraph(step)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

    # Closing note
    doc.add_paragraph("")  # one blank line before the note
    note = doc.add_paragraph(
        "Note: If none of the above steps resolve the issue, contact your IT Help Desk "
        "and provide the output of 'ipconfig /all' and the ping results for further diagnosis."
    )
    note.paragraph_format.space_before = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
