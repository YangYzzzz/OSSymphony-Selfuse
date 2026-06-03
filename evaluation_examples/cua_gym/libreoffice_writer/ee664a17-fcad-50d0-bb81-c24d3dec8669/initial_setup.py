"""
Initial Setup: Product catalog document with a 3-column table (no Discount column yet)
Task ID: writer_tbl_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_005'
OUTPUT = f'{WORKDIR}/product_catalog.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Add a title heading
    doc.add_heading('Product Catalog', level=1)

    # Add introductory paragraph
    doc.add_paragraph('The following table lists our current product offerings along with pricing and categories.')

    # Create table: 4 rows x 3 columns (header row + 3 data rows)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Row 0: Headers
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Item'
    header_cells[1].text = 'Price'
    header_cells[2].text = 'Category'

    # Make headers bold
    for cell in header_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True

    # Row 1: Laptop
    row1_cells = table.rows[1].cells
    row1_cells[0].text = 'Laptop'
    row1_cells[1].text = '$999'
    row1_cells[2].text = 'Electronics'

    # Row 2: Desk Chair
    row2_cells = table.rows[2].cells
    row2_cells[0].text = 'Desk Chair'
    row2_cells[1].text = '$249'
    row2_cells[2].text = 'Furniture'

    # Row 3: Headphones
    row3_cells = table.rows[3].cells
    row3_cells[0].text = 'Headphones'
    row3_cells[1].text = '$79'
    row3_cells[2].text = 'Electronics'

    # Add a closing paragraph
    doc.add_paragraph('Prices are subject to change without notice.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
