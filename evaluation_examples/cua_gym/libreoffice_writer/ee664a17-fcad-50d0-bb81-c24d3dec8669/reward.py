"""
Reward Script: Insert 'Discount' column to the right of 'Price' in product table
Task ID: writer_tbl_005
Domain: libreoffice_writer
Scoring:
  Component 1: Table has 4 columns (0.3 pts)
  Component 2: New column header at index 2 is 'Discount' (0.3 pts)
  Component 3: All 3 data rows have '10%' in the new Discount column (0.3 pts)
  Component 4: 'Category' column preserved at index 3 with correct values (0.1 pts)
"""

import os

from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_005'
FILE_PATH = f'{WORKDIR}/product_catalog.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — if it fails, we can't verify anything
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: document must have at least one table
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Component 1: Table has 4 columns (0.3 points)
    # Initial state has 3 columns; task requires inserting a new column → must be 4
    try:
        num_cols = len(table.columns)
        if num_cols == 4:
            print(f"PASS: Component 1 — Table has 4 columns (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — Expected 4 columns, found {num_cols}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: New column header at index 2 is 'Discount' (0.3 points)
    # Task specifies column must be inserted to the RIGHT of 'Price' (index 1),
    # so the new header must be at index 2, pushing 'Category' to index 3.
    try:
        header_row = table.rows[0]
        if len(header_row.cells) >= 3:
            header_text = header_row.cells[2].text.strip()
            if header_text == 'Discount':
                print(f"PASS: Component 2 — Header at column 2 is 'Discount' (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Expected 'Discount' at column 2, found '{header_text}'")
        else:
            print(f"FAIL: Component 2 — Header row has fewer than 3 cells")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All 3 data rows have '10%' in column index 2 (0.3 points)
    # Rows 1, 2, 3 should each contain '10%' in the Discount column (index 2)
    try:
        expected_discount = '10%'
        data_rows = table.rows[1:]  # skip header row
        discount_values = []
        for row in data_rows:
            if len(row.cells) >= 3:
                val = row.cells[2].text.strip()
                discount_values.append(val)
            else:
                discount_values.append(None)

        all_match = all(v == expected_discount for v in discount_values)
        if all_match and len(discount_values) == 3:
            print(f"PASS: Component 3 — All 3 data rows have '10%' in Discount column (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 — Discount column values: {discount_values}, expected ['10%', '10%', '10%']")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: 'Category' column preserved at index 3 with correct values (0.1 points)
    # After inserting the Discount column, 'Category' should shift from index 2 to index 3.
    # Original values: Electronics, Furniture, Electronics
    try:
        expected_header = 'Category'
        expected_categories = ['Electronics', 'Furniture', 'Electronics']

        header_cells = table.rows[0].cells
        if len(header_cells) >= 4:
            cat_header = header_cells[3].text.strip()
            cat_data = [table.rows[r].cells[3].text.strip() for r in range(1, 4)
                        if len(table.rows[r].cells) >= 4]

            if cat_header == expected_header and cat_data == expected_categories:
                print(f"PASS: Component 4 — 'Category' column preserved at index 3 with correct values (0.1 pts)")
                total_score += 0.1
            else:
                print(f"FAIL: Component 4 — Category header='{cat_header}', data={cat_data}")
        else:
            print(f"FAIL: Component 4 — Row has fewer than 4 cells")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
