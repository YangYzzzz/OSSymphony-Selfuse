"""
Initial Setup: Create a 5-page NDA contract document with no header.
Task ID: writer_legal_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup: standard letter
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Ensure NO header exists (default state, but be explicit)
    header = section.header
    header.is_linked_to_previous = True

    # ---- Title ----
    title = doc.add_heading('NON-DISCLOSURE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('NDA Agreement')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    # ---- Preamble ----
    doc.add_paragraph(
        'This Non-Disclosure Agreement ("Agreement") is entered into as of '
        'March 15, 2025, by and between Meridian Technologies, Inc., a Delaware '
        'corporation with its principal office at 4500 Innovation Drive, Suite 300, '
        'San Jose, CA 95134 ("Disclosing Party"), and Apex Consulting Group LLC, '
        'a California limited liability company with its principal office at '
        '1200 Market Street, Suite 800, San Francisco, CA 94103 ("Receiving Party").'
    )

    doc.add_paragraph(
        'WHEREAS, the Disclosing Party possesses certain confidential and proprietary '
        'information relating to its business operations, technology platforms, client '
        'relationships, and strategic initiatives; and'
    )

    doc.add_paragraph(
        'WHEREAS, the Receiving Party desires to receive certain confidential information '
        'for the purpose of evaluating a potential business relationship between the parties '
        '(the "Purpose");'
    )

    doc.add_paragraph(
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements contained '
        'herein, and for other good and valuable consideration, the receipt and sufficiency '
        'of which are hereby acknowledged, the parties agree as follows:'
    )

    # ---- Section 1 ----
    doc.add_heading('1. DEFINITION OF CONFIDENTIAL INFORMATION', level=1)

    doc.add_paragraph(
        '1.1 "Confidential Information" means any and all non-public information, whether '
        'written, oral, electronic, visual, or in any other form, disclosed by the '
        'Disclosing Party to the Receiving Party, including but not limited to:'
    )

    items_1 = [
        'Trade secrets, inventions, patent applications, and intellectual property;',
        'Business plans, financial data, projections, and strategic roadmaps;',
        'Customer lists, vendor agreements, pricing models, and sales methodologies;',
        'Software source code, algorithms, system architectures, and technical specifications;',
        'Marketing strategies, competitive analyses, and market research data;',
        'Personnel records, compensation structures, and organizational charts;',
        'Any information designated as "Confidential," "Proprietary," or with similar markings.',
    ]
    for item in items_1:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        '1.2 Confidential Information shall not include information that: (a) is or becomes '
        'publicly available through no fault of the Receiving Party; (b) was already in the '
        'possession of the Receiving Party prior to disclosure, as demonstrated by written records; '
        '(c) is independently developed by the Receiving Party without use of or reference to '
        'the Confidential Information; or (d) is rightfully obtained by the Receiving Party from '
        'a third party without restriction on disclosure.'
    )

    # ---- Section 2 ----
    doc.add_heading('2. OBLIGATIONS OF THE RECEIVING PARTY', level=1)

    doc.add_paragraph(
        '2.1 The Receiving Party agrees to hold all Confidential Information in strict '
        'confidence and to take all reasonable precautions to protect such information, '
        'using at least the same degree of care as the Receiving Party uses to protect '
        'its own confidential information, but in no event less than reasonable care.'
    )

    doc.add_paragraph(
        '2.2 The Receiving Party shall not, without the prior written consent of the '
        'Disclosing Party: (a) disclose any Confidential Information to any person or entity '
        'other than its employees, agents, or advisors who have a need to know such information '
        'for the Purpose and who are bound by obligations of confidentiality no less restrictive '
        'than those contained herein; (b) use any Confidential Information for any purpose other '
        'than the Purpose; or (c) reverse engineer, decompile, or disassemble any Confidential '
        'Information that constitutes software or technology.'
    )

    doc.add_paragraph(
        '2.3 The Receiving Party shall promptly notify the Disclosing Party in writing of any '
        'unauthorized use or disclosure of Confidential Information of which the Receiving Party '
        'becomes aware and shall cooperate with the Disclosing Party in every reasonable way to '
        'help the Disclosing Party regain possession of the Confidential Information and prevent '
        'its further unauthorized use or disclosure.'
    )

    # ---- Section 3 ----
    doc.add_heading('3. TERM AND TERMINATION', level=1)

    doc.add_paragraph(
        '3.1 This Agreement shall remain in effect for a period of three (3) years from the '
        'date first written above (the "Term"), unless earlier terminated by either party upon '
        'thirty (30) days written notice to the other party.'
    )

    doc.add_paragraph(
        '3.2 The obligations of confidentiality set forth in this Agreement shall survive '
        'the termination or expiration of this Agreement for a period of five (5) years '
        'following such termination or expiration, or for such longer period as may be '
        'required to protect trade secrets under applicable law.'
    )

    doc.add_paragraph(
        '3.3 Upon termination or expiration of this Agreement, or upon written request of '
        'the Disclosing Party, the Receiving Party shall promptly return or destroy all '
        'Confidential Information in its possession, including all copies, notes, summaries, '
        'and derivative works thereof, and shall certify in writing that it has done so.'
    )

    # ---- Section 4 ----
    doc.add_heading('4. REMEDIES', level=1)

    doc.add_paragraph(
        '4.1 The Receiving Party acknowledges and agrees that any breach or threatened breach '
        'of this Agreement may cause irreparable harm to the Disclosing Party for which monetary '
        'damages may be an inadequate remedy. Accordingly, the Disclosing Party shall be entitled '
        'to seek equitable relief, including injunction and specific performance, in addition to '
        'all other remedies available at law or in equity, without the necessity of proving actual '
        'damages or posting a bond or other security.'
    )

    doc.add_paragraph(
        '4.2 In the event that the Receiving Party is required by law, regulation, or legal '
        'process to disclose any Confidential Information, the Receiving Party shall, to the '
        'extent legally permissible, provide the Disclosing Party with prompt written notice of '
        'such requirement so that the Disclosing Party may seek a protective order or other '
        'appropriate remedy. The Receiving Party shall disclose only that portion of the '
        'Confidential Information which it is legally required to disclose and shall use '
        'reasonable efforts to obtain confidential treatment for such disclosed information.'
    )

    # ---- Section 5 ----
    doc.add_heading('5. INTELLECTUAL PROPERTY', level=1)

    doc.add_paragraph(
        '5.1 Nothing in this Agreement shall be construed as granting or conferring any rights '
        'by license or otherwise, expressly or impliedly, in any Confidential Information. '
        'All Confidential Information shall remain the sole and exclusive property of the '
        'Disclosing Party.'
    )

    doc.add_paragraph(
        '5.2 The Receiving Party shall not acquire any intellectual property rights in or to '
        'the Confidential Information by virtue of this Agreement or by virtue of access to '
        'or use of the Confidential Information, except as expressly set forth herein.'
    )

    # ---- Section 6 ----
    doc.add_heading('6. NO WARRANTY', level=1)

    doc.add_paragraph(
        '6.1 ALL CONFIDENTIAL INFORMATION IS PROVIDED "AS IS." THE DISCLOSING PARTY MAKES NO '
        'WARRANTIES, EXPRESS, IMPLIED, OR OTHERWISE, REGARDING THE ACCURACY, COMPLETENESS, OR '
        'PERFORMANCE OF ANY CONFIDENTIAL INFORMATION, AND HEREBY DISCLAIMS ALL WARRANTIES, '
        'INCLUDING THE IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE.'
    )

    # ---- Section 7 ----
    doc.add_heading('7. GENERAL PROVISIONS', level=1)

    doc.add_paragraph(
        '7.1 Governing Law. This Agreement shall be governed by and construed in accordance '
        'with the laws of the State of California, without regard to its conflict of laws principles.'
    )

    doc.add_paragraph(
        '7.2 Entire Agreement. This Agreement constitutes the entire agreement between the parties '
        'with respect to the subject matter hereof and supersedes all prior and contemporaneous '
        'agreements, understandings, negotiations, and discussions, whether oral or written, '
        'relating to such subject matter.'
    )

    doc.add_paragraph(
        '7.3 Amendment. This Agreement may not be amended, modified, or supplemented except by '
        'a written instrument signed by both parties.'
    )

    doc.add_paragraph(
        '7.4 Waiver. No waiver of any provision of this Agreement shall be effective unless in '
        'writing and signed by the party against whom the waiver is sought to be enforced. No '
        'failure or delay by either party in exercising any right, power, or privilege under this '
        'Agreement shall operate as a waiver thereof.'
    )

    doc.add_paragraph(
        '7.5 Severability. If any provision of this Agreement is held to be invalid or '
        'unenforceable by a court of competent jurisdiction, the remaining provisions shall '
        'continue in full force and effect, and the invalid or unenforceable provision shall be '
        'replaced by a valid and enforceable provision that most closely approximates the intent '
        'of the original provision.'
    )

    doc.add_paragraph(
        '7.6 Assignment. Neither party may assign or transfer this Agreement or any rights or '
        'obligations hereunder without the prior written consent of the other party, except that '
        'either party may assign this Agreement to a successor in connection with a merger, '
        'acquisition, or sale of all or substantially all of its assets.'
    )

    doc.add_paragraph(
        '7.7 Notices. All notices and other communications hereunder shall be in writing and '
        'shall be deemed duly given when delivered personally, sent by certified mail (return '
        'receipt requested), or sent by nationally recognized overnight courier service to the '
        'addresses set forth above or to such other address as either party may specify in writing.'
    )

    doc.add_paragraph(
        '7.8 Counterparts. This Agreement may be executed in counterparts, each of which shall '
        'be deemed an original, and all of which together shall constitute one and the same instrument.'
    )

    # ---- Signature Block ----
    doc.add_paragraph('')  # spacer
    doc.add_paragraph(
        'IN WITNESS WHEREOF, the parties have executed this Non-Disclosure Agreement as of '
        'the date first written above.'
    )

    doc.add_paragraph('')

    # Disclosing Party signature
    p1 = doc.add_paragraph()
    p1.add_run('DISCLOSING PARTY:').bold = True
    doc.add_paragraph('Meridian Technologies, Inc.')
    doc.add_paragraph('')
    doc.add_paragraph('By: ____________________________________')
    doc.add_paragraph('Name: Victoria R. Hammond')
    doc.add_paragraph('Title: Chief Executive Officer')
    doc.add_paragraph('Date: March 15, 2025')

    doc.add_paragraph('')

    # Receiving Party signature
    p2 = doc.add_paragraph()
    p2.add_run('RECEIVING PARTY:').bold = True
    doc.add_paragraph('Apex Consulting Group LLC')
    doc.add_paragraph('')
    doc.add_paragraph('By: ____________________________________')
    doc.add_paragraph('Name: David K. Nakamura')
    doc.add_paragraph('Title: Managing Director')
    doc.add_paragraph('Date: March 15, 2025')

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
