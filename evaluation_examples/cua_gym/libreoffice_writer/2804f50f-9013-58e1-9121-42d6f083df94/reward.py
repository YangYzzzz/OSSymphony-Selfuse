"""
Reward Script: Add a 'CONFIDENTIAL' header to every page of a contract document.
Task ID: writer_legal_005
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Header text contains 'CONFIDENTIAL'
  Component 2 (0.3): Header is explicitly defined (not linked to previous / empty)
  Component 3 (0.3): All sections carry the CONFIDENTIAL header (coverage check)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_005'


def persist_app_state(domain: str):
    """Best-effort save of any unsaved GUI edits."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that a 'CONFIDENTIAL' header has been added to the document.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)
    print(f"INFO: Document has {num_sections} section(s)")

    # Component 1: First section header text contains 'CONFIDENTIAL' (0.4 points)
    # This is the primary task requirement.
    try:
        section = doc.sections[0]
        header = section.header
        header_text = ""
        for para in header.paragraphs:
            header_text += para.text
        header_text_upper = header_text.strip().upper()
        if "CONFIDENTIAL" in header_text_upper:
            print(f"PASS: Component 1 — Header contains 'CONFIDENTIAL' (found: '{header_text.strip()}') (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Header text is '{header_text.strip()}', expected 'CONFIDENTIAL'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header is explicitly defined, not empty/linked-to-previous default (0.3 points)
    # In the initial file, header.is_linked_to_previous=True and text is empty.
    # In the golden file, header.is_linked_to_previous=False and text is set.
    try:
        section = doc.sections[0]
        header = section.header
        has_content = any(p.text.strip() for p in header.paragraphs)
        is_not_linked = not header.is_linked_to_previous
        if has_content and is_not_linked:
            print(f"PASS: Component 2 — Header is explicitly defined (is_linked_to_previous=False, has content) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Header not explicitly defined (is_linked_to_previous={header.is_linked_to_previous}, has_content={has_content})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All sections have 'CONFIDENTIAL' in header (0.3 points)
    # The task says "on every page", so all sections must carry the header.
    try:
        failing_sections = []
        for i, sec in enumerate(doc.sections):
            sec_header = sec.header
            sec_header_text = ""
            for para in sec_header.paragraphs:
                sec_header_text += para.text
            if "CONFIDENTIAL" not in sec_header_text.strip().upper():
                failing_sections.append((i, sec_header_text.strip()))
        if len(failing_sections) == 0:
            print(f"PASS: Component 3 — All {num_sections} section(s) have 'CONFIDENTIAL' in header (0.3 pts)")
            total_score += 0.3
        else:
            for idx, txt in failing_sections:
                print(f"FAIL: Component 3 — Section {idx} header does not contain 'CONFIDENTIAL' (found: '{txt}')")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
