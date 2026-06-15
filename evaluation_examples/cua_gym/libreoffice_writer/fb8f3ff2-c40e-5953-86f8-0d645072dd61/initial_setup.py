"""
Initial Setup: Create a #10 envelope document with return and delivery addresses.
Task ID: writer_lec_064
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_064'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Configure page as #10 envelope: 4.125" x 9.5"
    section = doc.sections[0]
    section.page_width = Inches(9.5)
    section.page_height = Inches(4.125)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # --- Return address (upper-left) ---
    return_addr_lines = [
        "Greenfield & Associates, LLC",
        "4720 Maple Ridge Drive, Suite 310",
        "Portland, OR 97205",
    ]
    for line in return_addr_lines:
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(10)

    # --- Spacing between return address and delivery address ---
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(36)
    spacer.paragraph_format.space_after = Pt(0)

    # --- Delivery address (centered, further down) ---
    delivery_lines = [
        "Ms. Rebecca Thornton",
        "Director of Operations",
        "Cascade Pacific Industries",
        "8901 Evergreen Boulevard",
        "Seattle, WA 98115",
    ]
    for line in delivery_lines:
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
