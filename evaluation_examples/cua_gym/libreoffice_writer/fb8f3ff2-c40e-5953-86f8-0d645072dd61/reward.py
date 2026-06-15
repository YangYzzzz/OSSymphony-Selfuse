"""
Reward Script: Add 'PRIORITY MAIL' in large red letters above the delivery address on an envelope
Task ID: writer_lec_064
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30) - 'PRIORITY MAIL' text exists in the document
  Component 2 (0.25) - Text is red colored (RGB close to FF0000)
  Component 3 (0.20) - Text is bold and large (>=18pt)
  Component 4 (0.25) - Text is positioned above the delivery address and addresses are intact
"""

import os
from math import sqrt
from docx import Document
from docx.shared import Pt, RGBColor

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_064'


def color_distance(c1, c2):
    """Euclidean distance between two RGB colors."""
    return sqrt((c1[0] - c2[0]) ** 2 + (c1[1] - c2[1]) ** 2 + (c1[2] - c2[2]) ** 2)


def find_priority_mail_para(doc):
    """Find paragraph containing 'PRIORITY MAIL' text. Returns (index, paragraph) or (None, None)."""
    for i, para in enumerate(doc.paragraphs):
        if 'PRIORITY MAIL' in para.text.upper().replace('  ', ' '):
            return i, para
    return None, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the PRIORITY MAIL paragraph
    pm_idx, pm_para = find_priority_mail_para(doc)

    # Component 1: 'PRIORITY MAIL' text exists (0.30 points)
    try:
        if pm_para is not None:
            print(f"PASS: Component 1 — 'PRIORITY MAIL' found at paragraph {pm_idx}: '{pm_para.text}' (0.30 pts)")
            total_score += 0.30
        else:
            print("FAIL: Component 1 — 'PRIORITY MAIL' text not found in any paragraph")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Text is red colored (0.25 points)
    try:
        if pm_para is not None:
            red_found = False
            for run in pm_para.runs:
                if run.text.strip() and run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    dist = color_distance(
                        (rgb[0], rgb[1], rgb[2]),
                        (0xFF, 0x00, 0x00)
                    )
                    if dist < 80:  # Allow some tolerance for red variants
                        red_found = True
                        print(f"PASS: Component 2 — Red color confirmed: RGB({rgb[0]},{rgb[1]},{rgb[2]}), distance={dist:.1f} (0.25 pts)")
                        total_score += 0.25
                        break
            if not red_found:
                colors = []
                for run in pm_para.runs:
                    if run.text.strip():
                        c = run.font.color.rgb if run.font.color and run.font.color.rgb else None
                        colors.append(str(c))
                print(f"FAIL: Component 2 — Expected red color, found: {colors}")
        else:
            print("FAIL: Component 2 — Skipped (no PRIORITY MAIL paragraph)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Text is bold and large (>=18pt) (0.20 points)
    try:
        if pm_para is not None:
            bold_ok = False
            size_ok = False
            actual_size = None
            for run in pm_para.runs:
                if run.text.strip():
                    if run.font.bold is True:
                        bold_ok = True
                    if run.font.size is not None and run.font.size.pt >= 18:
                        size_ok = True
                        actual_size = run.font.size.pt

            if bold_ok and size_ok:
                print(f"PASS: Component 3 — Bold=True, Size={actual_size}pt (>=18pt) (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 — Bold={bold_ok}, Size={'>=18pt' if size_ok else (str(actual_size) + 'pt' if actual_size else 'None')}")
        else:
            print("FAIL: Component 3 — Skipped (no PRIORITY MAIL paragraph)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: PRIORITY MAIL is positioned above delivery address AND addresses intact (0.25 points)
    try:
        if pm_para is not None:
            # Check that the delivery address is still present after PRIORITY MAIL
            # The delivery address starts with recipient name and ends with city/state/zip
            all_text = [p.text.strip() for p in doc.paragraphs]

            # Check return address is intact (first few paragraphs)
            return_addr_present = any('Greenfield' in t or 'greenfield' in t.lower() for t in all_text)
            # Check delivery address is intact
            delivery_addr_present = (
                any('Rebecca Thornton' in t for t in all_text) and
                any('Seattle' in t for t in all_text)
            )

            # Check PRIORITY MAIL appears before the delivery address recipient line
            delivery_idx = None
            for i, para in enumerate(doc.paragraphs):
                if 'Rebecca Thornton' in para.text:
                    delivery_idx = i
                    break

            positioned_correctly = (delivery_idx is not None and pm_idx < delivery_idx)

            if return_addr_present and delivery_addr_present and positioned_correctly:
                print(f"PASS: Component 4 — Addresses intact, PRIORITY MAIL at para {pm_idx} before delivery at para {delivery_idx} (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 4 — return_addr={return_addr_present}, delivery_addr={delivery_addr_present}, positioned={positioned_correctly}")
        else:
            print("FAIL: Component 4 — Skipped (no PRIORITY MAIL paragraph)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point — persist app state then verify
def persist_app_state(domain):
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
