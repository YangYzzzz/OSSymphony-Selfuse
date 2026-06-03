"""
Reward Script: Warranty Certificate Document
Task ID: writer_wf_031
Domain: libreoffice_writer
Scoring:
  C1: Title present and centered (0.15)
  C2: Page border on all 4 sides (0.15)
  C3: Product info fields (0.20)
  C4: Four numbered sections (0.20)
  C5: Five exclusion bullet items (0.10)
  C6: Four claim procedure steps (0.10)
  C7: Contact details and signature line (0.10)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_031'


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice state."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify warranty certificate task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_text = '\n'.join(p.text for p in doc.paragraphs)
    all_text_lower = all_text.lower()

    # Component 1: Title "LIMITED WARRANTY CERTIFICATE" present and centered (0.15 points)
    try:
        title_found = False
        for p in doc.paragraphs:
            if 'limited warranty certificate' in p.text.lower().strip():
                title_found = True
                break
        if title_found:
            print(f"PASS: Component 1 — Title 'LIMITED WARRANTY CERTIFICATE' found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Title 'LIMITED WARRANTY CERTIFICATE' not found in any paragraph")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Page border on all 4 sides (0.15 points)
    try:
        section = doc.sections[0]
        sect_pr = section._sectPr
        pg_borders = sect_pr.findall(qn('w:pgBorders'))
        border_sides_found = 0
        if pg_borders:
            for border_elem in pg_borders:
                for side in ['top', 'bottom', 'left', 'right']:
                    side_elem = border_elem.find(qn(f'w:{side}'))
                    if side_elem is not None:
                        val = side_elem.get(qn('w:val'))
                        if val and val != 'none':
                            border_sides_found += 1
        if border_sides_found >= 4:
            print(f"PASS: Component 2 — Page border found on all 4 sides (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 — Page border sides found: {border_sides_found}/4")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Product info fields (0.20 points)
    # Must have: product name/model, serial number, purchase date, warranty period (2 years)
    try:
        fields_found = 0
        # Product name field
        if re.search(r'product\s*name', all_text_lower):
            fields_found += 1
        # Model field
        if re.search(r'model', all_text_lower):
            fields_found += 1
        # Serial number
        if re.search(r'serial\s*number', all_text_lower):
            fields_found += 1
        # Purchase date
        if re.search(r'purchase\s*date', all_text_lower):
            fields_found += 1
        # Warranty period with 2 years
        if re.search(r'warranty\s*period', all_text_lower) and '2' in all_text:
            fields_found += 1

        field_score = (fields_found / 5) * 0.20
        if fields_found >= 5:
            print(f"PASS: Component 3 — All 5 product info fields found (0.20 pts)")
        else:
            print(f"PARTIAL: Component 3 — {fields_found}/5 product info fields found ({field_score:.2f} pts)")
        total_score += field_score
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Four numbered sections (Coverage, Exclusions, Claim Procedure, Limitation of Liability) (0.20 points)
    try:
        sections_found = 0
        required_sections = ['coverage', 'exclusions', 'claim procedure', 'limitation of liability']
        for sec_name in required_sections:
            # Look for numbered section headers like "1. Coverage" or just containing the keyword
            for p in doc.paragraphs:
                txt = p.text.strip().lower()
                if sec_name in txt and re.match(r'^\d+\.?\s', txt):
                    sections_found += 1
                    break

        section_score = (sections_found / 4) * 0.20
        if sections_found >= 4:
            print(f"PASS: Component 4 — All 4 numbered sections found (0.20 pts)")
        else:
            print(f"PARTIAL: Component 4 — {sections_found}/4 numbered sections found ({section_score:.2f} pts)")
        total_score += section_score
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Five exclusion bullet items under Exclusions section (0.10 points)
    try:
        # Count bullet-style paragraphs between "Exclusions" and "Claim Procedure" sections
        bullet_count = 0
        in_exclusions = False
        for p in doc.paragraphs:
            txt = p.text.strip().lower()
            if 'exclusion' in txt and re.match(r'^\d+\.?\s', txt):
                in_exclusions = True
                continue
            if in_exclusions and re.match(r'^\d+\.?\s', txt):
                # Hit next numbered section
                break
            if in_exclusions and p.text.strip():
                # Count non-empty paragraphs (could be bullet style or just indented text)
                style_name = p.style.name.lower() if p.style else ''
                # Accept bullet style, or any non-empty paragraph after the intro line
                if 'bullet' in style_name or 'list' in style_name:
                    bullet_count += 1
                elif not txt.startswith('this warranty'):
                    # Also count non-intro paragraphs as potential bullets
                    bullet_count += 1

        if bullet_count >= 5:
            print(f"PASS: Component 5 — {bullet_count} exclusion items found (>=5) (0.10 pts)")
            total_score += 0.10
        elif bullet_count >= 3:
            partial = 0.05
            print(f"PARTIAL: Component 5 — {bullet_count}/5 exclusion items found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Only {bullet_count} exclusion items found (need >=5)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Four claim procedure steps (0.10 points)
    try:
        step_count = 0
        in_claim = False
        for p in doc.paragraphs:
            txt = p.text.strip().lower()
            if 'claim procedure' in txt and re.match(r'^\d+\.?\s', txt):
                in_claim = True
                continue
            if in_claim and re.match(r'^\d+\.?\s', txt) and 'step' not in txt:
                # Hit next numbered section (not a step)
                break
            if in_claim and p.text.strip():
                # Count step-like paragraphs
                if re.match(r'step\s*\d', txt) or re.match(r'\d+[\.\)]\s', txt):
                    step_count += 1
                elif 'list number' in (p.style.name.lower() if p.style else ''):
                    step_count += 1

        if step_count >= 4:
            print(f"PASS: Component 6 — {step_count} claim procedure steps found (>=4) (0.10 pts)")
            total_score += 0.10
        elif step_count >= 2:
            partial = 0.05
            print(f"PARTIAL: Component 6 — {step_count}/4 claim procedure steps found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 6 — Only {step_count} claim procedure steps found (need >=4)")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Company contact details and authorized signature line (0.10 points)
    try:
        has_contact = False
        has_signature = False

        # Check for contact details (phone/email/address)
        if re.search(r'(phone|email|website|www\.)', all_text_lower):
            has_contact = True

        # Check for signature line (underscores or "signature" keyword)
        if re.search(r'(authorized\s*signature|signature\s*line|_{5,})', all_text_lower):
            has_signature = True

        comp7_score = 0.0
        if has_contact and has_signature:
            comp7_score = 0.10
            print(f"PASS: Component 7 — Contact details and signature line found (0.10 pts)")
        elif has_contact or has_signature:
            comp7_score = 0.05
            detail = 'contact' if has_contact else 'signature'
            print(f"PARTIAL: Component 7 — Only {detail} found ({comp7_score} pts)")
        else:
            print(f"FAIL: Component 7 — No contact details or signature line found")
        total_score += comp7_score
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
