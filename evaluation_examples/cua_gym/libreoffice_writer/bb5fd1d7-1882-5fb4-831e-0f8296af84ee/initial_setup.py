"""
Initial Setup: Legal memorandum without classification banners
Task ID: writer_legal_067
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_067'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title
    title = doc.add_heading('MEMORANDUM', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Memo header block
    memo_fields = [
        ('TO:', 'Sarah Mitchell, General Counsel'),
        ('FROM:', 'David Park, Senior Associate'),
        ('DATE:', 'March 28, 2026'),
        ('RE:', 'Analysis of Proposed Non-Compete Agreement — Westfield Technologies Acquisition'),
    ]
    for label, value in memo_fields:
        para = doc.add_paragraph()
        run_label = para.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_label.font.name = 'Times New Roman'
        run_value = para.add_run(f'\t{value}')
        run_value.font.size = Pt(11)
        run_value.font.name = 'Times New Roman'

    # Horizontal line
    doc.add_paragraph('_' * 72)

    # Section I
    h1 = doc.add_heading('I. Executive Summary', level=1)

    body1 = doc.add_paragraph()
    body1.paragraph_format.line_spacing = 1.5
    run = body1.add_run(
        'This memorandum addresses the enforceability of the proposed non-compete '
        'agreement in connection with the acquisition of Westfield Technologies, Inc. '
        'by Meridian Capital Partners, LLC. After reviewing applicable case law in the '
        'relevant jurisdictions (California, New York, and Delaware), we conclude that '
        'certain provisions of the draft agreement may face significant enforceability '
        'challenges, particularly under California Business and Professions Code '
        'Section 16600.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    body1b = doc.add_paragraph()
    body1b.paragraph_format.line_spacing = 1.5
    run = body1b.add_run(
        'Specifically, the five-year duration and the nationwide geographic scope '
        'of the proposed restrictions exceed what courts have historically deemed '
        'reasonable. We recommend narrowing the temporal and geographic limitations '
        'as described in Section IV below to improve the likelihood of enforcement.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # Section II
    doc.add_heading('II. Factual Background', level=1)

    body2 = doc.add_paragraph()
    body2.paragraph_format.line_spacing = 1.5
    run = body2.add_run(
        'Westfield Technologies, Inc. ("Westfield") is a Delaware corporation '
        'headquartered in San Jose, California, specializing in enterprise cloud '
        'infrastructure solutions. The company was founded in 2018 by Dr. Elena '
        'Vasquez and currently employs approximately 340 individuals across offices '
        'in San Jose, New York City, and Austin, Texas.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    body2b = doc.add_paragraph()
    body2b.paragraph_format.line_spacing = 1.5
    run = body2b.add_run(
        'Meridian Capital Partners, LLC ("Meridian") has proposed acquiring 100% '
        'of Westfield\'s outstanding shares for $485 million. As a condition of '
        'closing, Meridian requires that all key executives, including Dr. Vasquez, '
        'CTO James Morrison, and VP of Engineering Rachel Kim, execute non-compete '
        'agreements restricting their ability to engage in competing business '
        'activities for a period of five years following the closing date.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    body2c = doc.add_paragraph()
    body2c.paragraph_format.line_spacing = 1.5
    run = body2c.add_run(
        'The proposed non-compete agreement (the "Agreement") contains the following '
        'key provisions: (a) a five-year restriction period; (b) a nationwide '
        'geographic scope covering the entire United States; (c) a prohibition on '
        'employment with any entity engaged in cloud infrastructure, software-as-a-service, '
        'or related technology services; and (d) liquidated damages of $2.5 million '
        'per violation.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # Section III
    doc.add_heading('III. Legal Analysis', level=1)

    doc.add_heading('A. California Law', level=2)

    body3a = doc.add_paragraph()
    body3a.paragraph_format.line_spacing = 1.5
    run = body3a.add_run(
        'California Business and Professions Code Section 16600 provides that '
        '"every contract by which anyone is restrained from engaging in a lawful '
        'profession, trade, or business of any kind is to that extent void." '
        'The California Supreme Court in Edwards v. Arthur Andersen LLP (2008) '
        '44 Cal.4th 937 held that Section 16600 must be interpreted broadly, '
        'invalidating non-compete agreements except in narrow statutory exceptions '
        '(sale of a business under Section 16601, dissolution of a partnership '
        'under Section 16602, and dissolution of an LLC under Section 16602.5).'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    body3a2 = doc.add_paragraph()
    body3a2.paragraph_format.line_spacing = 1.5
    run = body3a2.add_run(
        'The acquisition of Westfield may qualify under the sale-of-business '
        'exception in Section 16601, which permits a seller to agree not to compete '
        'with the buyer in the same geographic area and line of business. However, '
        'the exception is narrowly construed, and the proposed five-year duration '
        'and nationwide scope may exceed the permissible boundaries even under this '
        'exception. See Strategix, Ltd. v. Infocure Corp. (2005) 142 Cal.App.4th 1068.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_heading('B. New York Law', level=2)

    body3b = doc.add_paragraph()
    body3b.paragraph_format.line_spacing = 1.5
    run = body3b.add_run(
        'New York courts apply a three-part test to evaluate the enforceability '
        'of restrictive covenants: (1) the restriction must be necessary to protect '
        'the employer\'s legitimate interests; (2) the restriction must not impose '
        'an undue hardship on the employee; and (3) the restriction must not be '
        'injurious to the public. See BDO Seidman v. Hirshberg, 93 N.Y.2d 382 (1999). '
        'New York courts have generally upheld non-compete agreements of one to two years '
        'but have shown skepticism toward restrictions exceeding three years.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_heading('C. Delaware Law', level=2)

    body3c = doc.add_paragraph()
    body3c.paragraph_format.line_spacing = 1.5
    run = body3c.add_run(
        'Delaware applies a reasonableness standard to non-compete agreements, '
        'examining the scope of the restriction in terms of time, geography, and '
        'activity. The Delaware Court of Chancery in Kodiak Building Partners, LLC '
        'v. Adams (2023) emphasized that restrictions must be narrowly tailored to '
        'protect legitimate business interests such as trade secrets, customer '
        'relationships, and goodwill. The court noted that overly broad restrictions '
        'that effectively prevent a former employee from working in their field '
        'entirely are unlikely to be enforced.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # Section IV
    doc.add_heading('IV. Recommendations', level=1)

    body4 = doc.add_paragraph()
    body4.paragraph_format.line_spacing = 1.5
    run = body4.add_run(
        'Based on the foregoing analysis, we recommend the following modifications '
        'to the proposed non-compete agreement to enhance enforceability across '
        'all relevant jurisdictions:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    recommendations = [
        'Reduce the restriction period from five years to two years from the closing date.',
        'Narrow the geographic scope to the metropolitan areas where Westfield currently '
        'operates (San Jose, New York City, and Austin), rather than the entire United States.',
        'Define "competing business" more precisely to include only companies deriving '
        'more than 30% of revenue from enterprise cloud infrastructure services.',
        'Reduce liquidated damages from $2.5 million to $500,000, with a provision for '
        'actual damages if they exceed the liquidated amount.',
        'Include a carve-out permitting key executives to serve as advisors or board '
        'members of non-competing technology companies.',
    ]
    for rec in recommendations:
        para = doc.add_paragraph(rec, style='List Number')
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # Section V
    doc.add_heading('V. Conclusion', level=1)

    body5 = doc.add_paragraph()
    body5.paragraph_format.line_spacing = 1.5
    run = body5.add_run(
        'The proposed non-compete agreement as currently drafted carries substantial '
        'risk of being found unenforceable, particularly in California. We strongly '
        'recommend adopting the modifications outlined above prior to presenting the '
        'Agreement to the key executives for signature. We are available to discuss '
        'these recommendations at your earliest convenience and to prepare revised '
        'draft language for your review.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
