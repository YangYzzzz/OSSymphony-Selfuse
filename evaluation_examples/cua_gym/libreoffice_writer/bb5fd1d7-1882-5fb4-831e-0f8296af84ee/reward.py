"""
Reward Script: Add classification banners (header + footer) to a legal document
Task ID: writer_legal_067
Domain: libreoffice_writer
Scoring:
  Component 1: Header text "ATTORNEY-CLIENT PRIVILEGED" present (0.20)
  Component 2: Header bold + red color (0.20)
  Component 3: Header centered alignment (0.10)
  Component 4: Header red bottom border (0.20)
  Component 5: Footer text "CONFIDENTIAL - DO NOT DISTRIBUTE" present (0.15)
  Component 6: Footer bold + red + centered (0.15)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_067'


def persist_app_state(domain):
    """Save any unsaved changes in LibreOffice Writer."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that classification banners have been added to header and footer.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Get the first section (banners should appear on every page)
    if len(doc.sections) == 0:
        print("FAIL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # === HEADER CHECKS ===

    # Component 1: Header contains "ATTORNEY-CLIENT PRIVILEGED" (0.20 points)
    try:
        header = section.header
        header_text = ""
        if header.paragraphs:
            header_text = " ".join(p.text.strip() for p in header.paragraphs).strip()

        if "ATTORNEY-CLIENT PRIVILEGED" in header_text.upper():
            print(f"PASS: Component 1 - Header contains 'ATTORNEY-CLIENT PRIVILEGED' (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 - Header text is {repr(header_text)}, expected 'ATTORNEY-CLIENT PRIVILEGED'")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Header text is bold and red (0.20 points)
    try:
        header = section.header
        found_bold_red = False
        for para in header.paragraphs:
            for run in para.runs:
                txt = run.text.strip().upper()
                if "ATTORNEY-CLIENT PRIVILEGED" in txt or (
                    "ATTORNEY" in txt and "PRIVILEGED" in txt
                ):
                    is_bold = run.font.bold is True
                    is_red = False
                    if run.font.color and run.font.color.rgb:
                        rgb = run.font.color.rgb
                        # RGBColor is an int-like; use str() to get hex
                        rgb_str = str(rgb).upper()
                        r_val = int(rgb_str[0:2], 16)
                        g_val = int(rgb_str[2:4], 16)
                        b_val = int(rgb_str[4:6], 16)
                        if r_val >= 200 and g_val <= 55 and b_val <= 55:
                            is_red = True
                    if is_bold and is_red:
                        found_bold_red = True
                        break
            if found_bold_red:
                break

        if found_bold_red:
            print(f"PASS: Component 2 - Header text is bold and red (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 - Header text is not bold+red")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Header is centered (0.10 points)
    try:
        header = section.header
        header_centered = False
        for para in header.paragraphs:
            if "ATTORNEY" in para.text.upper():
                al = para.paragraph_format.alignment
                if al == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    header_centered = True
                    break

        if header_centered:
            print(f"PASS: Component 3 - Header is centered (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 - Header is not centered")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Header has a red bottom border (0.20 points)
    try:
        header = section.header
        has_red_bottom_border = False
        for para in header.paragraphs:
            if "ATTORNEY" in para.text.upper():
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    pBdr = pPr.find(qn('w:pBdr'))
                    if pBdr is not None:
                        bottom = pBdr.find(qn('w:bottom'))
                        if bottom is not None:
                            bval = bottom.get(qn('w:val'))
                            bcolor = bottom.get(qn('w:color'))
                            # Border must exist (not 'none') and be red-ish
                            if bval and bval != 'none':
                                if bcolor:
                                    bcolor_upper = bcolor.upper()
                                    # Accept FF0000 or close red variants
                                    r = int(bcolor_upper[0:2], 16)
                                    g = int(bcolor_upper[2:4], 16)
                                    b = int(bcolor_upper[4:6], 16)
                                    if r >= 200 and g <= 55 and b <= 55:
                                        has_red_bottom_border = True

        if has_red_bottom_border:
            print(f"PASS: Component 4 - Header has red bottom border (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 - Header does not have a red bottom border")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # === FOOTER CHECKS ===

    # Component 5: Footer contains "CONFIDENTIAL - DO NOT DISTRIBUTE" (0.15 points)
    try:
        footer = section.footer
        footer_text = ""
        if footer.paragraphs:
            footer_text = " ".join(p.text.strip() for p in footer.paragraphs).strip()

        if "CONFIDENTIAL" in footer_text.upper() and "DO NOT DISTRIBUTE" in footer_text.upper():
            print(f"PASS: Component 5 - Footer contains 'CONFIDENTIAL - DO NOT DISTRIBUTE' (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 - Footer text is {repr(footer_text)}, expected 'CONFIDENTIAL - DO NOT DISTRIBUTE'")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: Footer text is bold, red, and centered (0.15 points)
    try:
        footer = section.footer
        found_bold_red_footer = False
        footer_centered = False

        for para in footer.paragraphs:
            if "CONFIDENTIAL" in para.text.upper():
                # Check centered
                al = para.paragraph_format.alignment
                if al == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    footer_centered = True

                # Check bold + red on runs
                for run in para.runs:
                    if "CONFIDENTIAL" in run.text.upper() or "DO NOT DISTRIBUTE" in run.text.upper():
                        is_bold = run.font.bold is True
                        is_red = False
                        if run.font.color and run.font.color.rgb:
                            rgb = run.font.color.rgb
                            rgb_str = str(rgb).upper()
                            r_val = int(rgb_str[0:2], 16)
                            g_val = int(rgb_str[2:4], 16)
                            b_val = int(rgb_str[4:6], 16)
                            if r_val >= 200 and g_val <= 55 and b_val <= 55:
                                is_red = True
                        if is_bold and is_red:
                            found_bold_red_footer = True
                            break

        if found_bold_red_footer and footer_centered:
            print(f"PASS: Component 6 - Footer is bold, red, centered (0.15 pts)")
            total_score += 0.15
        else:
            details = []
            if not found_bold_red_footer:
                details.append("not bold+red")
            if not footer_centered:
                details.append("not centered")
            print(f"FAIL: Component 6 - Footer formatting issues: {', '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
