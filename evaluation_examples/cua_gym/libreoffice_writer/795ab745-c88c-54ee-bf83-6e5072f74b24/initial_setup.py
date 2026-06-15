"""
Initial Setup: Mail merge label document without Next Record fields
Task ID: writer_mt_026
Domain: libreoffice_writer

Creates a label document with a 3x10 table grid per page (3 pages, 90 labels total).
Each label cell contains merge fields (Name, Address, City, State, ZIP) but is
MISSING the 'Next Record' field — that's the task the agent must complete.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_026'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

ROWS_PER_PAGE = 10
COLS = 3
PAGES = 3


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_merge_field(paragraph, field_name):
    """Add a MERGEFIELD field code to a paragraph."""
    # Field structure: fldChar(begin) + instrText + fldChar(separate) + display text + fldChar(end)
    run_begin = paragraph.add_run()
    fld_begin = run_begin._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run_begin._element.append(fld_begin)

    run_instr = paragraph.add_run()
    instr = run_instr._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr.text = f' MERGEFIELD {field_name} '
    run_instr._element.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = run_sep._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run_sep._element.append(fld_sep)

    # Display text (placeholder shown before merge)
    run_display = paragraph.add_run(f'\u00AB{field_name}\u00BB')
    run_display.font.size = Pt(9)

    run_end = paragraph.add_run()
    fld_end = run_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run_end._element.append(fld_end)


def populate_label_cell(cell):
    """Populate a label cell with merge fields for a shipping label."""
    # Clear default paragraph
    cell.paragraphs[0].text = ''

    # Name line
    p_name = cell.paragraphs[0]
    p_name.paragraph_format.space_after = Pt(0)
    p_name.paragraph_format.space_before = Pt(2)
    add_merge_field(p_name, 'Name')

    # Address line
    p_addr = cell.add_paragraph()
    p_addr.paragraph_format.space_after = Pt(0)
    p_addr.paragraph_format.space_before = Pt(0)
    add_merge_field(p_addr, 'Address')

    # City, State ZIP line
    p_csz = cell.add_paragraph()
    p_csz.paragraph_format.space_after = Pt(2)
    p_csz.paragraph_format.space_before = Pt(0)
    add_merge_field(p_csz, 'City')
    run_comma = p_csz.add_run(', ')
    run_comma.font.size = Pt(9)
    add_merge_field(p_csz, 'State')
    run_space = p_csz.add_run(' ')
    run_space.font.size = Pt(9)
    add_merge_field(p_csz, 'ZIP')

    # Set cell margins for label appearance
    tc_pr = cell._element.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="72" w:type="dxa"/>'
        f'  <w:left w:w="115" w:type="dxa"/>'
        f'  <w:bottom w:w="72" w:type="dxa"/>'
        f'  <w:right w:w="115" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(margins)


def create_initial():
    doc = Document()

    # Page setup: US Letter, narrow margins for labels
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.19)
    section.right_margin = Inches(0.19)

    # Title paragraph
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run_title = title.add_run('Shipping Labels — Mail Merge Template')
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.name = 'Arial'

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    run_sub = subtitle.add_run('Data source: ShippingAddresses (90 records)')
    run_sub.font.size = Pt(9)
    run_sub.font.name = 'Arial'
    run_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for page_idx in range(PAGES):
        # Create the 10x3 label table for this page
        table = doc.add_table(rows=ROWS_PER_PAGE, cols=COLS)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Set column widths (~2.7 inches each)
        col_width = Inches(2.7)
        for col in table.columns:
            for cell in col.cells:
                cell.width = col_width

        # Set row heights
        for row in table.rows:
            row.height = Inches(1.0)

        # Populate each cell with merge fields
        for r in range(ROWS_PER_PAGE):
            for c in range(COLS):
                cell = table.cell(r, c)
                populate_label_cell(cell)

        # Add page break between pages (not after last)
        if page_idx < PAGES - 1:
            doc.add_page_break()

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
