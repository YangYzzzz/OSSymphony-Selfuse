"""
Reward Script: Add 'Next Record' fields to mail merge label document
Task ID: writer_mt_026
Domain: libreoffice_writer
Scoring:
  Component 1 (0.50): Non-first cells in each table have a NEXT field (87 of 90 cells)
  Component 2 (0.30): Total NEXT field count is exactly 87 (29 per table x 3 tables)
  Component 3 (0.20): Correct pattern — first cells lack NEXT AND non-first cells have NEXT
                       (compound check: both parts must hold for points)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_026'


def count_next_fields_in_cell(cell):
    """Count NEXT (not NEXTIF) instrText fields in a cell's XML."""
    xml = cell._element.xml
    # Match instrText containing NEXT but not NEXTIF
    instrs = re.findall(r'<w:instrText[^>]*>(.*?)</w:instrText>', xml)
    count = 0
    for instr in instrs:
        stripped = instr.strip()
        if stripped == 'NEXT' or stripped.startswith('NEXT '):
            # Exclude NEXTIF
            if not stripped.startswith('NEXTIF'):
                count += 1
    return count


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have exactly 3 tables (3 pages of labels)
    if len(doc.tables) != 3:
        print(f"PRECONDITION FAIL: Expected 3 tables, found {len(doc.tables)}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: each table must be 10 rows x 3 cols
    for ti, table in enumerate(doc.tables):
        if len(table.rows) != 10 or len(table.columns) != 3:
            print(f"PRECONDITION FAIL: Table {ti} is {len(table.rows)}x{len(table.columns)}, expected 10x3")
            print("REWARD: 0.0")
            return 0.0

    # Component 1: Non-first cells have NEXT field (0.50 points)
    # In golden, all cells except [0,0] in each table should have a NEXT field
    # That's 29 cells per table = 87 total. This FAILS on initial (0 have NEXT).
    try:
        non_first_with_next = 0
        non_first_total = 0
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    if ri == 0 and ci == 0:
                        continue  # skip first cell of each table
                    non_first_total += 1
                    if count_next_fields_in_cell(cell) >= 1:
                        non_first_with_next += 1

        ratio = non_first_with_next / non_first_total if non_first_total > 0 else 0
        if ratio >= 0.95:
            print(f"PASS: Component 1 — {non_first_with_next}/{non_first_total} non-first cells have NEXT field (0.50 pts)")
            total_score += 0.50
        elif ratio > 0:
            partial = 0.50 * ratio
            print(f"PARTIAL: Component 1 — {non_first_with_next}/{non_first_total} non-first cells have NEXT field ({partial:.3f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — {non_first_with_next}/{non_first_total} non-first cells have NEXT field")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Total NEXT field count is exactly 87 (0.30 points)
    # 3 tables x 29 non-first cells = 87 total. FAILS on initial (count is 0).
    try:
        total_next = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_next += count_next_fields_in_cell(cell)

        if total_next == 87:
            print(f"PASS: Component 2 — Total NEXT fields = {total_next} (expected 87) (0.30 pts)")
            total_score += 0.30
        elif 80 <= total_next <= 90:
            closeness = 1.0 - abs(total_next - 87) / 87
            partial = 0.30 * closeness
            print(f"PARTIAL: Component 2 — Total NEXT fields = {total_next} (expected 87) ({partial:.3f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Total NEXT fields = {total_next} (expected 87)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Correct pattern — first cells lack NEXT AND at least 85 non-first cells
    # have NEXT (0.20 points). This is a compound check that verifies the correct
    # distribution pattern. FAILS on initial because non-first cells have 0 NEXT fields.
    try:
        first_cells_no_next = 0
        for ti, table in enumerate(doc.tables):
            first_cell = table.rows[0].cells[0]
            if count_next_fields_in_cell(first_cell) == 0:
                first_cells_no_next += 1

        # Both conditions must hold: first cells correct AND non-first cells have NEXT
        if first_cells_no_next == 3 and non_first_with_next >= 85:
            print(f"PASS: Component 3 — Correct pattern: {first_cells_no_next}/3 first cells lack NEXT, "
                  f"{non_first_with_next}/87 non-first cells have NEXT (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — Pattern check: first_cells_no_next={first_cells_no_next}/3, "
                  f"non_first_with_next={non_first_with_next}/87 (need both correct)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
