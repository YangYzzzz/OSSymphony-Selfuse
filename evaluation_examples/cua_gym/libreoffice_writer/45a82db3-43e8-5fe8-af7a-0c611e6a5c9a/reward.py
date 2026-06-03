"""
Reward Script: Format heading with Heading 1 style, Arial Black 24pt, center, dark blue (#003366)
Task ID: writer_rd_002
Domain: libreoffice_writer
Scoring:
  Component 1: Heading 1 style applied (0.2)
  Component 2: Font name Arial Black (0.2)
  Component 3: Font size 24pt (0.2)
  Component 4: Font color #003366 (0.2)
  Component 5: Center alignment (0.2)
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_002'


def persist_app_state(domain: str):
    """Try to save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError as e:
        print(f"CRITICAL: Missing library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has at least 1 paragraph with 'Annual Report 2025'
    if len(doc.paragraphs) < 1 or 'Annual Report 2025' not in doc.paragraphs[0].text:
        print(f"PRECONDITION FAIL: First paragraph text is not 'Annual Report 2025'")
        print("REWARD: 0.0")
        return 0.0

    heading_para = doc.paragraphs[0]
    pf = heading_para.paragraph_format

    # Component 1: Heading 1 style applied (0.2 points)
    try:
        style_name = heading_para.style.name
        if style_name == 'Heading 1':
            print(f"PASS: Component 1 — Style is 'Heading 1' (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 1 — Expected style 'Heading 1', found '{style_name}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Font name is Arial Black (0.2 points)
    try:
        runs = heading_para.runs
        if not runs:
            print(f"FAIL: Component 2 — No runs found in heading paragraph")
        elif all(r.font.name == 'Arial Black' for r in runs):
            print(f"PASS: Component 2 — Font is 'Arial Black' (0.2 pts)")
            total_score += 0.2
        else:
            bad = [r for r in runs if r.font.name != 'Arial Black']
            print(f"FAIL: Component 2 — Expected font 'Arial Black', found '{bad[0].font.name}'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Font size is 24pt (0.2 points)
    try:
        runs = heading_para.runs
        if not runs:
            print(f"FAIL: Component 3 — No runs found")
        elif all(r.font.size is not None and r.font.size.pt == 24.0 for r in runs):
            print(f"PASS: Component 3 — Font size is 24pt (0.2 pts)")
            total_score += 0.2
        else:
            for r in runs:
                sz = r.font.size.pt if r.font.size else None
                if sz != 24.0:
                    print(f"FAIL: Component 3 — Expected 24pt, found {sz}pt")
                    break
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Font color is #003366 dark blue (0.2 points)
    try:
        runs = heading_para.runs
        target_color = RGBColor(0x00, 0x33, 0x66)
        if not runs:
            print(f"FAIL: Component 4 — No runs found")
        elif all(r.font.color.rgb is not None and str(r.font.color.rgb).upper() == str(target_color).upper() for r in runs):
            print(f"PASS: Component 4 — Font color is #003366 (0.2 pts)")
            total_score += 0.2
        else:
            for r in runs:
                rgb = r.font.color.rgb
                if rgb is None or str(rgb).upper() != str(target_color).upper():
                    print(f"FAIL: Component 4 — Expected color #003366, found {rgb}")
                    break
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Center alignment (0.2 points)
    try:
        alignment = pf.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            print(f"PASS: Component 5 — Alignment is CENTER (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 5 — Expected CENTER alignment, found {alignment}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
