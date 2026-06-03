"""
Initial Setup: Format heading with Heading 1 style, Arial Black 24pt, center, dark blue
Task ID: writer_rd_002
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_002'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: heading text with Default Paragraph Style (no heading style, no formatting)
    para_heading = doc.add_paragraph('Annual Report 2025')
    # Ensure it is plain default style — no special formatting
    run = para_heading.runs[0]
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)

    # Paragraph 2: body text
    para2 = doc.add_paragraph(
        'The fiscal year 2025 has been a transformative period for Meridian '
        'Technologies. Revenue grew by 18.3% year-over-year, reaching $247.5 '
        'million, driven primarily by expansion into the Asia-Pacific market and '
        'the successful launch of our CloudSync platform in Q2.'
    )

    # Paragraph 3: body text
    para3 = doc.add_paragraph(
        'Our research and development division invested $34.2 million in next-generation '
        'artificial intelligence solutions, resulting in three patent filings and the '
        'release of SmartAssist 3.0, which has already been adopted by over 1,200 '
        'enterprise clients across North America and Europe.'
    )

    # Paragraph 4: body text
    para4 = doc.add_paragraph(
        'Looking ahead, the board has approved a strategic acquisition of DataPulse '
        'Analytics for $89 million, expected to close in Q1 2026. This acquisition '
        'will strengthen our data infrastructure capabilities and position Meridian '
        'as a leader in the predictive analytics market segment.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
