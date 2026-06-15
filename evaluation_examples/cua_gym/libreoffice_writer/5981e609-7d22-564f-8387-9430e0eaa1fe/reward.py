"""
Reward Script: Find & Replace DD/MM/YYYY dates with MM-DD-YYYY format
Task ID: osworld_writer_find_replace_004
Domain: libreoffice_writer
Scoring:
  Component 1: All 6 DD/MM/YYYY dates removed from document    (0.5 pts)
  Component 2: All 6 MM-DD-YYYY dates correctly present        (0.5 pts)
  Total: 1.0
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_find_replace_004'

# Ground truth: 6 dates that must be converted from DD/MM/YYYY -> MM-DD-YYYY
# Initial format: DD/MM/YYYY
# Expected format after conversion: MM-DD-YYYY
EXPECTED_CONVERTED_DATES = [
    '02-01-2024',  # was 01/02/2024
    '03-15-2024',  # was 15/03/2024
    '04-20-2024',  # was 20/04/2024
    '06-30-2024',  # was 30/06/2024
    '09-01-2024',  # was 01/09/2024
    '12-31-2024',  # was 31/12/2024
]

ORIGINAL_DATES = [
    '01/02/2024',
    '15/03/2024',
    '20/04/2024',
    '30/06/2024',
    '01/09/2024',
    '31/12/2024',
]


def get_all_text(doc):
    """Extract all text from document paragraphs and tables."""
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    texts.append(para.text)
    return ' '.join(texts)


def verify_task(file_path):
    """
    Verify that the Find & Replace task was completed:
    - All 6 DD/MM/YYYY dates have been replaced with MM-DD-YYYY format.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document (precondition gate)
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    full_text = get_all_text(doc)

    # Component 1: No original DD/MM/YYYY dates remain (0.5 points)
    # This checks that all 6 original dates have been removed/replaced.
    # FAILS on initial_env (6 DD/MM/YYYY dates present), PASSES on golden_env (0 remain).
    try:
        dd_mm_yyyy_matches = re.findall(r'\b(\d{2}/\d{2}/\d{4})\b', full_text)
        remaining_original = [d for d in dd_mm_yyyy_matches if d in ORIGINAL_DATES]
        num_remaining = len(remaining_original)

        if num_remaining == 0:
            print(f"PASS: Component 1 — No DD/MM/YYYY dates remain (all 6 converted) (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — {num_remaining} DD/MM/YYYY date(s) still present: {remaining_original}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All 6 expected MM-DD-YYYY dates are present (0.5 points)
    # This checks that the correct converted dates appear in the document.
    # FAILS on initial_env (no MM-DD-YYYY dates), PASSES on golden_env (all 6 present).
    try:
        mm_dd_yyyy_matches = re.findall(r'\b(\d{2}-\d{2}-\d{4})\b', full_text)
        found_expected = [d for d in EXPECTED_CONVERTED_DATES if d in mm_dd_yyyy_matches]
        missing_dates = [d for d in EXPECTED_CONVERTED_DATES if d not in mm_dd_yyyy_matches]
        num_found = len(found_expected)

        if num_found == 6:
            print(f"PASS: Component 2 — All 6 MM-DD-YYYY dates present: {found_expected} (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 2 — Only {num_found}/6 expected dates found. "
                  f"Missing: {missing_dates}. Found: {found_expected}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
