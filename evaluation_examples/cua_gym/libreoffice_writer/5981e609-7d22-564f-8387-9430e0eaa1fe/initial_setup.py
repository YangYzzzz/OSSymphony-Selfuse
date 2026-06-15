"""
Initial Setup: Legal contract document with 6 dates in DD/MM/YYYY format
Task ID: osworld_writer_find_replace_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_find_replace_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('SERVICE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph with date 1: 01/02/2024
    doc.add_paragraph(
        'This Service Agreement ("Agreement") is entered into as of 01/02/2024 '
        'by and between Meridian Consulting Group LLC, a Delaware limited liability '
        'company ("Service Provider"), and Hartwell Industries Inc., a California '
        'corporation ("Client").'
    )

    # Section 1
    doc.add_heading('1. TERM OF AGREEMENT', level=1)
    doc.add_paragraph(
        'The term of this Agreement shall commence on 15/03/2024 and shall continue '
        'for a period of twelve (12) months unless earlier terminated in accordance '
        'with the provisions set forth herein. Either party may terminate this '
        'Agreement upon thirty (30) days written notice to the other party.'
    )

    # Section 2
    doc.add_heading('2. SERVICES', level=1)
    doc.add_paragraph(
        'Service Provider agrees to perform the following consulting services for '
        'the Client during the term of this Agreement: strategic planning, market '
        'analysis, operational efficiency review, and executive advisory services. '
        'All services shall be delivered in accordance with the project timeline '
        'agreed upon by both parties.'
    )

    # Section 3 with date 3: 20/04/2024
    doc.add_heading('3. COMPENSATION', level=1)
    doc.add_paragraph(
        'Client agrees to pay Service Provider a monthly retainer fee of $18,500 '
        'due on the first business day of each month. The first payment shall be '
        'due on 20/04/2024. All payments shall be made by wire transfer to the '
        'bank account specified by Service Provider in writing.'
    )

    # Section 4
    doc.add_heading('4. CONFIDENTIALITY', level=1)
    doc.add_paragraph(
        'Each party acknowledges that in connection with this Agreement it may '
        'receive certain confidential and proprietary information and materials '
        'of the other party. Each party agrees to hold the other party\'s '
        'Confidential Information in strict confidence and not to disclose '
        'such information to any third parties without prior written consent.'
    )

    # Section 5 with date 4: 30/06/2024
    doc.add_heading('5. INTELLECTUAL PROPERTY', level=1)
    doc.add_paragraph(
        'All work product, deliverables, and materials created by Service Provider '
        'in the performance of this Agreement shall become the exclusive property '
        'of Client upon full payment. Transfer of intellectual property rights '
        'shall be effective as of 30/06/2024, provided all outstanding invoices '
        'have been settled in full by that date.'
    )

    # Section 6
    doc.add_heading('6. LIMITATION OF LIABILITY', level=1)
    doc.add_paragraph(
        'In no event shall Service Provider be liable for any indirect, incidental, '
        'special, consequential, or punitive damages arising out of or related to '
        'this Agreement. Service Provider\'s total cumulative liability shall not '
        'exceed the total fees paid by Client in the three months preceding the '
        'event giving rise to liability.'
    )

    # Section 7 with date 5: 01/09/2024
    doc.add_heading('7. REVIEW AND RENEWAL', level=1)
    doc.add_paragraph(
        'A comprehensive performance review shall be conducted on 01/09/2024 to '
        'assess the progress of all ongoing projects and the overall satisfaction '
        'of both parties. Based on the outcome of this review, the parties may '
        'mutually agree to extend, modify, or terminate this Agreement prior to '
        'its scheduled expiration date.'
    )

    # Section 8
    doc.add_heading('8. DISPUTE RESOLUTION', level=1)
    doc.add_paragraph(
        'Any dispute arising out of or relating to this Agreement shall first be '
        'addressed through good-faith negotiation between the parties. If the '
        'dispute cannot be resolved within thirty (30) days of written notice, '
        'the parties agree to submit the matter to binding arbitration conducted '
        'in accordance with the rules of the American Arbitration Association.'
    )

    # Section 9 with date 6: 31/12/2024
    doc.add_heading('9. EXPIRATION AND TERMINATION', level=1)
    doc.add_paragraph(
        'Unless renewed in writing by both parties, this Agreement shall '
        'automatically expire on 31/12/2024. Any renewal must be executed at '
        'least thirty (30) days prior to the expiration date. Upon expiration '
        'or termination, Service Provider shall promptly return all Client '
        'materials and cease use of any Client confidential information.'
    )

    # Section 10
    doc.add_heading('10. GOVERNING LAW', level=1)
    doc.add_paragraph(
        'This Agreement shall be governed by and construed in accordance with '
        'the laws of the State of Delaware, without regard to its conflict of '
        'law provisions. The parties consent to the exclusive jurisdiction of '
        'the courts located in Wilmington, Delaware for any legal proceedings '
        'arising under this Agreement.'
    )

    # Signature block
    doc.add_heading('SIGNATURES', level=1)
    doc.add_paragraph(
        'IN WITNESS WHEREOF, the parties hereto have executed this Agreement '
        'as of the date first written above.'
    )

    # Signature table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    # Row 0: party names
    table.cell(0, 0).text = 'MERIDIAN CONSULTING GROUP LLC'
    table.cell(0, 1).text = 'HARTWELL INDUSTRIES INC.'

    # Row 1: signature lines
    table.cell(1, 0).text = 'Signature: _______________________'
    table.cell(1, 1).text = 'Signature: _______________________'

    # Row 2: name lines
    table.cell(2, 0).text = 'Name: Jonathan R. Mercer'
    table.cell(2, 1).text = 'Name: Diana L. Hartwell'

    # Row 3: title lines
    table.cell(3, 0).text = 'Title: Chief Executive Officer'
    table.cell(3, 1).text = 'Title: President & CEO'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
