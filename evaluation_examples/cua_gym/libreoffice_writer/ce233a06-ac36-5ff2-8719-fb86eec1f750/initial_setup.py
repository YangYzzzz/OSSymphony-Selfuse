"""
Initial Setup: Create an offer letter with left-aligned company name
Task ID: writer_hr_008
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Company Name (LEFT-aligned, this is what the task asks to center) ---
    company_para = doc.add_paragraph()
    company_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    company_para.paragraph_format.space_after = Pt(4)
    run = company_para.add_run('Meridian Technologies Inc.')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # --- Tagline ---
    tagline = doc.add_paragraph()
    tagline.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    tagline.paragraph_format.space_after = Pt(18)
    tag_run = tagline.add_run('Innovating the Future of Enterprise Solutions')
    tag_run.italic = True
    tag_run.font.size = Pt(10)
    tag_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Horizontal rule via bottom border on empty paragraph ---
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(12)
    date_run = date_para.add_run('March 18, 2025')
    date_run.font.size = Pt(11)

    # --- Recipient ---
    recipient_lines = [
        'Ms. Elena Vasquez',
        '742 Brookfield Avenue, Apt 3B',
        'Portland, OR 97205'
    ]
    for line in recipient_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(11)

    # --- Spacing ---
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

    # --- Subject ---
    subject = doc.add_paragraph()
    subject.paragraph_format.space_after = Pt(12)
    subj_run = subject.add_run('Re: Offer of Employment - Senior Data Engineer')
    subj_run.bold = True
    subj_run.font.size = Pt(11)

    # --- Greeting ---
    greeting = doc.add_paragraph()
    greeting.paragraph_format.space_after = Pt(8)
    gr_run = greeting.add_run('Dear Ms. Vasquez,')
    gr_run.font.size = Pt(11)

    # --- Body paragraphs ---
    body_texts = [
        'We are pleased to extend this formal offer of employment for the position of '
        'Senior Data Engineer at Meridian Technologies Inc. After a thorough evaluation '
        'of your qualifications and interview performance, we are confident that your '
        'expertise in distributed systems and real-time data pipelines will be a valuable '
        'asset to our Data Platform team.',

        'The terms of your employment are outlined below:',
    ]
    for text in body_texts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.font.size = Pt(11)

    # --- Terms table ---
    terms = [
        ('Position', 'Senior Data Engineer'),
        ('Department', 'Data Platform Engineering'),
        ('Reporting To', 'Dr. Raj Patel, VP of Engineering'),
        ('Start Date', 'April 14, 2025'),
        ('Location', 'Portland, OR (Hybrid - 3 days on-site)'),
        ('Annual Base Salary', '$142,000'),
        ('Sign-On Bonus', '$15,000 (paid within 30 days of start)'),
        ('Annual Bonus Target', '12% of base salary'),
        ('Equity Grant', '2,500 RSUs vesting over 4 years'),
        ('PTO', '22 days annually + 10 company holidays'),
        ('Benefits', 'Medical, dental, vision, 401(k) with 4% match'),
    ]

    table = doc.add_table(rows=len(terms), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(terms):
        cell0 = table.cell(i, 0)
        cell0_run = cell0.paragraphs[0].add_run(label)
        cell0_run.bold = True
        cell0_run.font.size = Pt(10)
        cell1 = table.cell(i, 1)
        cell1_run = cell1.paragraphs[0].add_run(value)
        cell1_run.font.size = Pt(10)

    # --- Post-table paragraphs ---
    spacer2 = doc.add_paragraph()
    spacer2.paragraph_format.space_after = Pt(4)

    closing_texts = [
        'This offer is contingent upon the successful completion of a background check '
        'and verification of your educational credentials. Please note that employment '
        'at Meridian Technologies Inc. is at-will, meaning either party may terminate '
        'the employment relationship at any time, with or without cause or notice.',

        'To accept this offer, please sign and return this letter by April 1, 2025. '
        'Should you have any questions regarding the terms of employment or the onboarding '
        'process, please do not hesitate to contact our HR department at '
        'hr@meridiantech.com or (503) 555-0147.',

        'We are excited about the prospect of having you join our team and look forward '
        'to the contributions you will bring to Meridian Technologies.',
    ]
    for text in closing_texts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.font.size = Pt(11)

    # --- Signature block ---
    spacer3 = doc.add_paragraph()
    spacer3.paragraph_format.space_after = Pt(20)

    sig_line = doc.add_paragraph()
    sig_run = sig_line.add_run('Warm regards,')
    sig_run.font.size = Pt(11)

    spacer4 = doc.add_paragraph()
    spacer4.paragraph_format.space_after = Pt(24)

    name_line = doc.add_paragraph()
    name_run = name_line.add_run('Catherine Whitmore')
    name_run.bold = True
    name_run.font.size = Pt(11)

    title_line = doc.add_paragraph()
    title_line.paragraph_format.space_after = Pt(0)
    title_run = title_line.add_run('Director of Human Resources')
    title_run.font.size = Pt(11)

    company_line = doc.add_paragraph()
    company_line.paragraph_format.space_after = Pt(20)
    comp_run = company_line.add_run('Meridian Technologies Inc.')
    comp_run.font.size = Pt(11)

    # --- Acceptance block ---
    accept_header = doc.add_paragraph()
    accept_header.paragraph_format.space_after = Pt(12)
    ah_run = accept_header.add_run('ACCEPTANCE')
    ah_run.bold = True
    ah_run.font.size = Pt(12)

    accept_text = doc.add_paragraph()
    accept_text.paragraph_format.space_after = Pt(24)
    at_run = accept_text.add_run(
        'I, Elena Vasquez, hereby accept the offer of employment as described above.'
    )
    at_run.font.size = Pt(11)

    sig_fields = [
        'Signature: ____________________________    Date: ______________',
    ]
    for field in sig_fields:
        p = doc.add_paragraph()
        r = p.add_run(field)
        r.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
