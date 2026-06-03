"""
Reward Script: Center-align the company name 'Meridian Technologies Inc.' at the top of the offer letter.
Task ID: writer_hr_008
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): First paragraph contains 'Meridian Technologies Inc.' AND is center-aligned
  Component 2 (0.4): Alignment verified via XML jVal attribute as CENTER (enum=1 or 'center')
"""

import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_008'


def persist_app_state(domain: str):
    """Attempt to save any unsaved GUI state."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 1 paragraph
    if len(doc.paragraphs) == 0:
        print("FAIL: Document has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    # Find the paragraph containing 'Meridian Technologies Inc.'
    target_para = None
    target_idx = None
    for i, para in enumerate(doc.paragraphs):
        if 'Meridian Technologies Inc.' in para.text:
            target_para = para
            target_idx = i
            break

    if target_para is None:
        print("FAIL: Could not find paragraph containing 'Meridian Technologies Inc.'")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Found target paragraph at index {target_idx}: \"{target_para.text}\"")

    # Component 1: Target paragraph is center-aligned via python-docx API (0.6 points)
    # This FAILS on initial (LEFT) and PASSES on golden (CENTER)
    try:
        alignment = target_para.paragraph_format.alignment
        print(f"INFO: Paragraph alignment value: {alignment}")
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            print(f"PASS: Component 1 — 'Meridian Technologies Inc.' is center-aligned (0.6 pts)")
            total_score += 0.6
        else:
            align_name = str(alignment) if alignment is not None else "None (inherited/LEFT)"
            print(f"FAIL: Component 1 — Expected CENTER, found: {align_name}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Verify center alignment via raw XML as independent check (0.4 points)
    # This FAILS on initial (no jc='center') and PASSES on golden (jc='center')
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        pPr = target_para._element.find('.//w:pPr', ns)
        if pPr is not None:
            jc = pPr.find('w:jc', ns)
            if jc is not None:
                jc_val = jc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                print(f"INFO: XML jc val: '{jc_val}'")
                if jc_val == 'center':
                    print(f"PASS: Component 2 — XML confirms center alignment (0.4 pts)")
                    total_score += 0.4
                else:
                    print(f"FAIL: Component 2 — XML jc val is '{jc_val}', expected 'center'")
            else:
                print("FAIL: Component 2 — No jc element found in paragraph properties (not center)")
        else:
            print("FAIL: Component 2 — No paragraph properties element found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
