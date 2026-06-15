"""
Initial Setup: Create framed_note.docx with a text box on page 1 (10cm x 4cm) with thin default borders
Task ID: writer_obj_048
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
DESKTOP = '/home/user/Desktop'
TASK_ID = 'writer_obj_048'
OUTPUT = f'{DESKTOP}/framed_note.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set page to standard letter
    section = doc.sections[0]
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Add a heading and some paragraph text before the text box
    doc.add_heading("Project Status Update", level=1)
    doc.add_paragraph(
        "This document provides a summary of the current project status, "
        "key milestones, and upcoming deliverables for the Q2 planning cycle."
    )

    # Add a text box using direct XML (python-docx doesn't have direct textbox API)
    # Text box dimensions: 10cm wide x 4cm tall
    # Using default (thin) border: 1pt (12700 EMU) solid black
    # Default padding: 91440 EMU (0.1 inches = ~2.54mm)

    # Width: 10cm = 3600000 EMU, Height: 4cm = 1440000 EMU
    width_emu = int(10 * 360000)   # 10cm in EMU
    height_emu = int(4 * 360000)   # 4cm in EMU

    # Default internal margins: ~91440 EMU (approximately 0.1 inch = ~7.2pt = 2.54mm)
    # For "default" we use 91440 (Word default is 0.1 inch)
    inset = 91440  # default padding ~0.1in = 7.2pt

    # Default thin border: 1pt = 12700 EMU, color black (000000)
    border_sz = "8"     # 1pt in half-points (8 * 0.5pt = 4pt? No: in w:sz, unit is 1/8 pt; 8 = 1pt)
    # Actually in drawingML, line width is in EMU. For a thin default: use 9525 EMU (0.75pt)
    # but for simplicity, use w:sz="6" (6/8 = 0.75pt) -- thin border in OOXML

    # Build the drawing XML for a floating text box
    drawing_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
         xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
         xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
         xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
         xmlns:v="urn:schemas-microsoft-com:vml"
         xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
  <w:r>
    <w:rPr/>
    <w:drawing>
      <wp:anchor distT="114300" distB="114300" distL="114300" distR="114300"
                 simplePos="0" relativeHeight="251658240" behindDoc="0"
                 locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column">
          <wp:align>left</wp:align>
        </wp:positionH>
        <wp:positionV relativeFrom="paragraph">
          <wp:posOffset>0</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapSquare wrapText="bothSides"/>
        <wp:docPr id="1" name="TextBox 1"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp>
              <wps:cNvSpPr txBx="1">
                <a:spLocks noChangeArrowheads="1"/>
              </wps:cNvSpPr>
              <wps:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{width_emu}" cy="{height_emu}"/>
                </a:xfrm>
                <a:prstGeom prst="rect">
                  <a:avLst/>
                </a:prstGeom>
                <a:noFill/>
                <a:ln w="9525">
                  <a:solidFill>
                    <a:srgbClr val="000000"/>
                  </a:solidFill>
                </a:ln>
              </wps:spPr>
              <wps:txbx>
                <w:txbxContent>
                  <w:p>
                    <w:r>
                      <w:t xml:space="preserve">Important Notice: Please review all project deliverables by the end of this quarter. Ensure all team members have submitted their status reports and updated task completion records in the shared project tracker.</w:t>
                    </w:r>
                  </w:p>
                </w:txbxContent>
              </wps:txbx>
              <wps:bodyPr insTwpEmu="{inset}" insBtwEmu="{inset}" insLEmu="{inset}" insREmu="{inset}"
                          rot="0" spcFirstLastPara="0" vertOverflow="overflow" horzOverflow="overflow"
                          vert="horz" wrap="square" lIns="{inset}" tIns="{inset}" rIns="{inset}" bIns="{inset}"
                          numCol="1" spcCol="0" rtlCol="0" fromWordArt="0" anchor="t" anchorCtr="0"
                          forceAA="0" compatLnSpc="1">
                <a:prstTxWarp prst="textNoShape">
                  <a:avLst/>
                </a:prstTxWarp>
              </wps:bodyPr>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </w:r>
</w:p>'''

    # Parse and insert the drawing paragraph after the second body paragraph
    drawing_elem = etree.fromstring(drawing_xml)
    doc.element.body.append(drawing_elem)

    # Add more content below
    doc.add_paragraph("")
    doc.add_paragraph(
        "The following sections outline current progress on each workstream. "
        "Please direct any questions or concerns to the project lead."
    )

    para2 = doc.add_paragraph()
    run2 = para2.add_run("Key Deliverables:")
    run2.bold = True

    items = [
        "Finalize requirements documentation by March 15, 2026",
        "Complete integration testing for Module A and Module B",
        "Submit budget forecast to finance team",
        "Schedule stakeholder review meeting",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
