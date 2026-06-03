"""
Initial Setup: User guide document with 8 chapters using Heading 1 style
Task ID: writer_rd_048
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_048'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # ---- Configure Heading 1 style: Liberation Sans 18pt bold black, no borders ----
    style_h1 = doc.styles['Heading 1']
    font = style_h1.font
    font.name = 'Liberation Sans'
    font.size = Pt(18)
    font.bold = True
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # black

    # Ensure no bottom border on the style
    pf = style_h1.paragraph_format
    pf.space_after = Pt(6)  # small default, NOT 1.0 cm
    pf.space_before = Pt(12)

    # Set rFonts for the style to ensure Liberation Sans renders
    rpr = style_h1.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:ascii'), 'Liberation Sans')
    rfonts.set(qn('w:hAnsi'), 'Liberation Sans')
    rfonts.set(qn('w:cs'), 'Liberation Sans')

    # Remove any borders that might exist on Heading 1 paragraph style
    pPr = style_h1.element.get_or_add_pPr()
    for old_bdr in pPr.findall(qn('w:pBdr')):
        pPr.remove(old_bdr)

    # ---- Body style ----
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Liberation Sans'
    style_normal.font.size = Pt(11)

    # ---- Document Title ----
    title_para = doc.add_heading('CloudSync Platform - User Guide', level=0)

    intro = doc.add_paragraph(
        'This comprehensive user guide covers the installation, configuration, '
        'and daily operation of the CloudSync Platform. It is designed for system '
        'administrators, IT managers, and end users who need to manage cloud-based '
        'file synchronization across enterprise environments.'
    )

    doc.add_paragraph(
        'Version 3.2.1 | Last Updated: March 2026 | Confidential'
    )

    # ---- Chapter 1 ----
    doc.add_heading('Chapter 1: Getting Started', level=1)
    doc.add_paragraph(
        'Before installing CloudSync Platform, verify that your server meets the minimum '
        'hardware requirements: 8 GB RAM, 4-core CPU, and 500 GB SSD storage. The software '
        'supports Ubuntu 22.04 LTS, Red Hat Enterprise Linux 9, and Windows Server 2022.'
    )
    doc.add_paragraph(
        'Download the latest installer from the CloudSync Portal at '
        'https://portal.cloudsync-platform.io/downloads. Choose the package that matches '
        'your operating system and architecture. For clustered deployments, download the '
        'Enterprise bundle which includes load-balancer configuration templates.'
    )
    doc.add_paragraph(
        'After downloading, verify the checksum using SHA-256 to ensure file integrity. '
        'Run the installer with elevated privileges and follow the on-screen prompts to '
        'complete the base installation.'
    )

    # ---- Chapter 2 ----
    doc.add_heading('Chapter 2: System Configuration', level=1)
    doc.add_paragraph(
        'The main configuration file is located at /etc/cloudsync/config.yaml. Key parameters '
        'include sync_interval (default: 300 seconds), max_concurrent_uploads (default: 10), '
        'and encryption_mode (AES-256-GCM recommended for enterprise deployments).'
    )
    doc.add_paragraph(
        'Database connectivity is configured in the [database] section. CloudSync supports '
        'PostgreSQL 14+, MySQL 8.0+, and MariaDB 10.6+. Connection pooling is enabled by '
        'default with a pool size of 20 connections. Adjust pool_max_size based on your '
        'expected concurrent user load.'
    )

    # ---- Chapter 3 ----
    doc.add_heading('Chapter 3: User Management', level=1)
    doc.add_paragraph(
        'CloudSync provides role-based access control (RBAC) with four predefined roles: '
        'Administrator, Manager, Editor, and Viewer. Each role can be customized through the '
        'web administration panel at https://admin.cloudsync-platform.io.'
    )
    doc.add_paragraph(
        'To create a new user account, navigate to Settings > User Management > Add User. '
        'Fill in the required fields including email address, department, and assigned role. '
        'The system will send an activation email with a temporary password that expires '
        'after 72 hours.'
    )
    doc.add_paragraph(
        'Multi-factor authentication (MFA) can be enforced at the organization level. '
        'Supported methods include TOTP authenticator apps, SMS verification, and hardware '
        'security keys (FIDO2/WebAuthn).'
    )

    # ---- Chapter 4 ----
    doc.add_heading('Chapter 4: File Synchronization', level=1)
    doc.add_paragraph(
        'The synchronization engine uses a delta-sync algorithm that transfers only modified '
        'blocks of data, reducing bandwidth consumption by up to 85% compared to full-file '
        'transfers. Conflict resolution follows a last-writer-wins policy by default, but '
        'can be configured for manual resolution in collaborative environments.'
    )
    doc.add_paragraph(
        'Selective sync allows users to choose specific folders for synchronization. This is '
        'particularly useful for remote workers with limited bandwidth or storage. Configure '
        'selective sync rules in the client application under Preferences > Sync Settings.'
    )

    # ---- Chapter 5 ----
    doc.add_heading('Chapter 5: Security and Compliance', level=1)
    doc.add_paragraph(
        'All data in transit is encrypted using TLS 1.3. Data at rest is encrypted with '
        'AES-256-GCM. Encryption keys are managed through an integrated key management '
        'system that supports automatic key rotation on a configurable schedule (default: '
        '90 days).'
    )
    doc.add_paragraph(
        'CloudSync is compliant with SOC 2 Type II, ISO 27001, GDPR, and HIPAA. Audit '
        'logs capture all file access, modification, and sharing events with tamper-proof '
        'timestamps. Logs are retained for 7 years by default and can be exported to SIEM '
        'solutions via Syslog or REST API.'
    )

    # ---- Chapter 6 ----
    doc.add_heading('Chapter 6: Monitoring and Alerts', level=1)
    doc.add_paragraph(
        'The built-in monitoring dashboard provides real-time visibility into sync status, '
        'storage utilization, active sessions, and error rates. Custom dashboards can be '
        'created using the drag-and-drop widget editor.'
    )
    doc.add_paragraph(
        'Alert rules are configured in Settings > Notifications > Alert Rules. Common alerts '
        'include: storage quota exceeded (threshold: 90%), sync failure rate above 5%, '
        'unauthorized access attempts, and certificate expiration warnings (30 days before). '
        'Notifications can be sent via email, Slack, Microsoft Teams, or PagerDuty.'
    )

    # ---- Chapter 7 ----
    doc.add_heading('Chapter 7: Backup and Recovery', level=1)
    doc.add_paragraph(
        'CloudSync performs incremental backups every 4 hours with full backups scheduled '
        'weekly on Sunday at 02:00 UTC. Backup destinations include local NAS, Amazon S3, '
        'Azure Blob Storage, and Google Cloud Storage. Cross-region replication ensures '
        'data durability with a 99.999999999% (11 nines) guarantee.'
    )
    doc.add_paragraph(
        'To restore files, navigate to the Recovery Console and select the desired restore '
        'point. Point-in-time recovery allows you to restore to any moment within the '
        'retention period (default: 90 days). Granular recovery supports restoring individual '
        'files, folders, or entire user accounts.'
    )

    # ---- Chapter 8 ----
    doc.add_heading('Chapter 8: Troubleshooting', level=1)
    doc.add_paragraph(
        'Common issues and resolutions are documented in the CloudSync Knowledge Base at '
        'https://support.cloudsync-platform.io/kb. Before contacting support, collect the '
        'following diagnostic information: server logs from /var/log/cloudsync/, client logs '
        'from ~/.cloudsync/logs/, and a screenshot of any error messages.'
    )
    doc.add_paragraph(
        'For sync conflicts, check the Conflict Resolution Queue under Tools > Conflicts. '
        'Each entry shows the conflicting file versions, modification timestamps, and the '
        'users involved. Resolve conflicts by choosing the preferred version or merging '
        'changes manually.'
    )
    doc.add_paragraph(
        'If the server becomes unresponsive, restart the CloudSync service using: '
        'systemctl restart cloudsync-server. Check the status with: systemctl status '
        'cloudsync-server. For persistent issues, increase log verbosity by setting '
        'log_level=DEBUG in config.yaml and reproduce the issue.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
