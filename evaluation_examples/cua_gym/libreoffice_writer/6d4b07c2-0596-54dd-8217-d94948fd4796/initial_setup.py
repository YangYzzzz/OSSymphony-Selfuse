"""
Initial Setup: Technical manual with headings and hyperlinks for PDF export task
Task ID: writer_tech_056
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_056'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = paragraph._element.makeelement(qn('w:hyperlink'), {
        qn('r:id'): r_id,
    })
    new_run = paragraph._element.makeelement(qn('w:r'), {})
    rPr = paragraph._element.makeelement(qn('w:rPr'), {})
    rStyle = paragraph._element.makeelement(qn('w:rStyle'), {qn('w:val'): 'Hyperlink'})
    color = paragraph._element.makeelement(qn('w:color'), {qn('w:val'): '0563C1'})
    u = paragraph._element.makeelement(qn('w:u'), {qn('w:val'): 'single'})
    rPr.append(rStyle)
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)
    t = paragraph._element.makeelement(qn('w:t'), {})
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading('CloudSync Platform Technical Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Version 3.2.1 — April 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('Prepared by the CloudSync Engineering Team')
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph('')  # spacer

    # ============================================================
    # Heading 1: Introduction
    # ============================================================
    doc.add_heading('1. Introduction', level=1)

    doc.add_heading('1.1 Purpose', level=2)
    p = doc.add_paragraph(
        'This manual provides comprehensive technical documentation for the CloudSync Platform, '
        'a distributed file synchronization and collaboration system designed for enterprise deployments. '
        'It covers architecture, installation, configuration, API reference, and troubleshooting procedures.'
    )

    doc.add_heading('1.2 Audience', level=2)
    doc.add_paragraph(
        'This document is intended for system administrators, DevOps engineers, and software developers '
        'who are responsible for deploying, configuring, and maintaining CloudSync Platform instances.'
    )

    doc.add_heading('1.3 Related Resources', level=2)
    p = doc.add_paragraph('For additional information, refer to the following online resources:')
    doc.add_paragraph('CloudSync Developer Portal', style='List Bullet')
    p_link1 = doc.add_paragraph('  ')
    add_hyperlink(p_link1, 'https://developer.cloudsync.example.com', 'https://developer.cloudsync.example.com')

    doc.add_paragraph('CloudSync API Reference', style='List Bullet')
    p_link2 = doc.add_paragraph('  ')
    add_hyperlink(p_link2, 'https://api.cloudsync.example.com/docs', 'https://api.cloudsync.example.com/docs')

    doc.add_paragraph('Community Support Forum', style='List Bullet')
    p_link3 = doc.add_paragraph('  ')
    add_hyperlink(p_link3, 'https://community.cloudsync.example.com', 'https://community.cloudsync.example.com')

    # ============================================================
    # Heading 1: System Architecture
    # ============================================================
    doc.add_heading('2. System Architecture', level=1)

    doc.add_heading('2.1 High-Level Overview', level=2)
    doc.add_paragraph(
        'CloudSync employs a microservices architecture with the following core components: '
        'the Sync Engine, the Metadata Service, the Storage Abstraction Layer, and the Authentication Gateway. '
        'These components communicate via gRPC and event-driven messaging through Apache Kafka.'
    )

    doc.add_heading('2.1.1 Sync Engine', level=3)
    doc.add_paragraph(
        'The Sync Engine is the central component responsible for detecting file changes, computing deltas, '
        'and orchestrating bidirectional synchronization between client devices and the cloud storage backend. '
        'It uses a Merkle tree-based approach for efficient change detection across large directory hierarchies.'
    )

    doc.add_heading('2.1.2 Metadata Service', level=3)
    doc.add_paragraph(
        'The Metadata Service maintains a versioned index of all files and directories managed by CloudSync. '
        'It stores file attributes, version history, sharing permissions, and conflict resolution state in a '
        'PostgreSQL database with read replicas for horizontal scalability.'
    )

    doc.add_heading('2.1.3 Storage Abstraction Layer', level=3)
    doc.add_paragraph(
        'The Storage Abstraction Layer provides a unified interface over multiple backend storage systems, '
        'including AWS S3, Azure Blob Storage, Google Cloud Storage, and on-premises MinIO clusters. '
        'This allows operators to choose or migrate storage backends without modifying application logic.'
    )

    doc.add_heading('2.2 Network Topology', level=2)
    doc.add_paragraph(
        'In a production deployment, CloudSync components are distributed across three availability zones. '
        'The load balancer routes client traffic to the nearest API gateway, which forwards requests '
        'to the appropriate microservice. Inter-service communication uses mutual TLS (mTLS) for encryption.'
    )

    # ============================================================
    # Heading 1: Installation Guide
    # ============================================================
    doc.add_heading('3. Installation Guide', level=1)

    doc.add_heading('3.1 System Requirements', level=2)

    doc.add_heading('3.1.1 Hardware Requirements', level=3)
    doc.add_paragraph('Minimum hardware requirements for a single-node deployment:')
    doc.add_paragraph('CPU: 8 cores (Intel Xeon E5 or equivalent)', style='List Bullet')
    doc.add_paragraph('RAM: 32 GB DDR4 ECC', style='List Bullet')
    doc.add_paragraph('Storage: 500 GB NVMe SSD for metadata, plus object storage backend', style='List Bullet')
    doc.add_paragraph('Network: 10 Gbps Ethernet', style='List Bullet')

    doc.add_heading('3.1.2 Software Prerequisites', level=3)
    doc.add_paragraph('The following software must be installed before deploying CloudSync:')
    doc.add_paragraph('Ubuntu Server 22.04 LTS or RHEL 9.x', style='List Bullet')
    doc.add_paragraph('Docker Engine 24.0+ with Docker Compose v2', style='List Bullet')
    doc.add_paragraph('PostgreSQL 15+ (for metadata service)', style='List Bullet')
    doc.add_paragraph('Apache Kafka 3.5+ (for event messaging)', style='List Bullet')
    doc.add_paragraph('Redis 7.0+ (for session caching)', style='List Bullet')

    doc.add_heading('3.2 Installation Steps', level=2)
    doc.add_paragraph(
        'Follow these steps to install CloudSync on a fresh Ubuntu 22.04 server. '
        'For RHEL-based installations, consult the platform-specific appendix.'
    )
    doc.add_paragraph('Clone the CloudSync deployment repository from GitHub:', style='List Number')
    p_git = doc.add_paragraph('  ')
    add_hyperlink(p_git, 'https://github.com/cloudsync/deploy', 'https://github.com/cloudsync/deploy')

    doc.add_paragraph('Run the prerequisite installation script: ./scripts/install-deps.sh', style='List Number')
    doc.add_paragraph('Configure environment variables in .env (see Section 4.1)', style='List Number')
    doc.add_paragraph('Execute docker compose up -d to start all services', style='List Number')
    doc.add_paragraph('Verify the deployment using the health check endpoint (Section 3.3)', style='List Number')

    doc.add_heading('3.3 Post-Installation Verification', level=2)
    doc.add_paragraph(
        'After installation, verify that all services are running correctly by accessing the health endpoint. '
        'Navigate to https://your-server:8443/api/v1/health in a browser or use curl. '
        'A healthy deployment returns HTTP 200 with a JSON response listing all service statuses.'
    )

    # ============================================================
    # Heading 1: Configuration Reference
    # ============================================================
    doc.add_heading('4. Configuration Reference', level=1)

    doc.add_heading('4.1 Environment Variables', level=2)
    doc.add_paragraph(
        'CloudSync is configured primarily through environment variables defined in the .env file. '
        'The following table lists the most commonly used configuration parameters.'
    )

    # Configuration table
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    headers = ['Variable', 'Default', 'Description']
    for i, h in enumerate(headers):
        run = table.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True

    config_data = [
        ['CS_DB_HOST', 'localhost', 'PostgreSQL server hostname'],
        ['CS_DB_PORT', '5432', 'PostgreSQL server port'],
        ['CS_KAFKA_BROKERS', 'localhost:9092', 'Comma-separated Kafka broker addresses'],
        ['CS_REDIS_URL', 'redis://localhost:6379', 'Redis connection URL for session cache'],
        ['CS_STORAGE_BACKEND', 's3', 'Storage backend type: s3, azure, gcs, minio'],
        ['CS_LOG_LEVEL', 'info', 'Logging level: debug, info, warn, error'],
        ['CS_MAX_FILE_SIZE', '5368709120', 'Maximum file size in bytes (default 5 GB)'],
    ]
    for r, row_data in enumerate(config_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_heading('4.2 Security Configuration', level=2)
    doc.add_paragraph(
        'CloudSync supports OAuth 2.0, SAML 2.0, and LDAP for user authentication. '
        'TLS certificates must be provided for all public-facing endpoints. '
        'For certificate management, see the official documentation:'
    )
    p_cert = doc.add_paragraph()
    add_hyperlink(p_cert, 'https://docs.cloudsync.example.com/security/tls', 'https://docs.cloudsync.example.com/security/tls')

    doc.add_heading('4.2.1 OAuth 2.0 Integration', level=3)
    doc.add_paragraph(
        'To configure OAuth 2.0 authentication, register CloudSync as a client application with your '
        'identity provider (e.g., Okta, Azure AD, or Keycloak). Set the CS_OAUTH_CLIENT_ID and '
        'CS_OAUTH_CLIENT_SECRET environment variables with the credentials obtained during registration.'
    )

    doc.add_heading('4.2.2 Encryption at Rest', level=3)
    doc.add_paragraph(
        'CloudSync supports AES-256 encryption for data at rest. Enable it by setting CS_ENCRYPTION_ENABLED=true '
        'and providing a 256-bit encryption key via CS_ENCRYPTION_KEY. Key rotation is supported through the '
        'administrative API endpoint /api/v1/admin/rotate-key.'
    )

    # ============================================================
    # Heading 1: API Reference
    # ============================================================
    doc.add_heading('5. API Reference', level=1)

    doc.add_heading('5.1 REST API Endpoints', level=2)
    doc.add_paragraph(
        'The CloudSync REST API is available at https://your-server:8443/api/v1/. '
        'All endpoints require authentication via Bearer token in the Authorization header. '
        'Rate limiting is enforced at 1000 requests per minute per authenticated user.'
    )

    doc.add_heading('5.1.1 File Operations', level=3)
    doc.add_paragraph(
        'POST /api/v1/files/upload — Upload a new file or update an existing file. '
        'Supports multipart form data with a maximum payload size configured by CS_MAX_FILE_SIZE. '
        'Returns a JSON response containing the file ID, version number, and upload timestamp.'
    )
    doc.add_paragraph(
        'GET /api/v1/files/{id}/download — Download a specific file by its unique identifier. '
        'Supports range requests for partial downloads and resumable transfers. '
        'The response includes Content-Disposition and ETag headers for caching.'
    )

    doc.add_heading('5.1.2 Synchronization Endpoints', level=3)
    doc.add_paragraph(
        'POST /api/v1/sync/delta — Submit a batch of local changes for server-side conflict resolution. '
        'The request body contains a JSON array of change records, each specifying the file path, '
        'operation type (create, modify, delete, move), and a content hash for integrity verification.'
    )

    doc.add_heading('5.2 Webhook Notifications', level=2)
    doc.add_paragraph(
        'CloudSync can send real-time webhook notifications when files are modified, shared, or deleted. '
        'Configure webhook endpoints through the administration dashboard or via the API. '
        'For webhook payload formats, refer to:'
    )
    p_webhook = doc.add_paragraph()
    add_hyperlink(p_webhook, 'https://api.cloudsync.example.com/docs/webhooks', 'https://api.cloudsync.example.com/docs/webhooks')

    # ============================================================
    # Heading 1: Troubleshooting
    # ============================================================
    doc.add_heading('6. Troubleshooting', level=1)

    doc.add_heading('6.1 Common Issues', level=2)

    doc.add_heading('6.1.1 Sync Conflicts', level=3)
    doc.add_paragraph(
        'When two users modify the same file simultaneously, CloudSync creates a conflict copy. '
        'Conflict copies are named with the pattern "filename (conflict - username - timestamp).ext". '
        'Users can resolve conflicts through the web interface or the desktop client.'
    )

    doc.add_heading('6.1.2 Performance Degradation', level=3)
    doc.add_paragraph(
        'If synchronization performance degrades over time, check the following: '
        'database query performance (pg_stat_statements), Kafka consumer lag (kafka-consumer-groups.sh), '
        'and Redis memory usage. A full troubleshooting guide is available at:'
    )
    p_trouble = doc.add_paragraph()
    add_hyperlink(p_trouble, 'https://docs.cloudsync.example.com/troubleshooting', 'https://docs.cloudsync.example.com/troubleshooting')

    doc.add_heading('6.2 Log Analysis', level=2)
    doc.add_paragraph(
        'CloudSync writes structured JSON logs to /var/log/cloudsync/. Use the cs-logquery tool '
        'to search and filter logs by service, severity, and time range. For centralized logging, '
        'configure the CS_LOG_EXPORT environment variable to forward logs to your ELK stack or Datadog instance.'
    )

    doc.add_heading('6.3 Support Channels', level=2)
    doc.add_paragraph('If you need further assistance, contact the CloudSync support team:')
    doc.add_paragraph('Email: support@cloudsync.example.com', style='List Bullet')
    p_support = doc.add_paragraph()
    p_support.style = doc.styles['List Bullet']
    p_support.add_run('Support Portal: ')
    add_hyperlink(p_support, 'https://support.cloudsync.example.com', 'https://support.cloudsync.example.com')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
