"""
Initial Setup: Table with 5 rows and 4 columns, no formatting
Task ID: writer_tbl_048
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_048'
OUTPUT = f'{WORKDIR}/highlight_table.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Add a title paragraph
    title = doc.add_paragraph("Project Budget Overview")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    # Add a short introductory paragraph
    doc.add_paragraph(
        "The following table summarizes active projects, their team leads, and allocated budgets for the current fiscal year."
    )

    # Add the table: 5 rows x 4 columns
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"

    # Row 1 (header): ID | Project | Lead | Budget
    header_data = ["ID", "Project", "Lead", "Budget"]
    header_row = table.rows[0]
    for col_idx, text in enumerate(header_data):
        cell = header_row.cells[col_idx]
        # Clear default empty paragraph and set text
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.bold = True

    # Row 2: 001 | Alpha | Kim | $50K
    row2_data = ["001", "Alpha", "Kim", "$50K"]
    row2 = table.rows[1]
    for col_idx, text in enumerate(row2_data):
        cell = row2.cells[col_idx]
        cell.paragraphs[0].clear()
        cell.paragraphs[0].add_run(text)

    # Row 3: 002 | Beta | Lee | $75K
    row3_data = ["002", "Beta", "Lee", "$75K"]
    row3 = table.rows[2]
    for col_idx, text in enumerate(row3_data):
        cell = row3.cells[col_idx]
        cell.paragraphs[0].clear()
        cell.paragraphs[0].add_run(text)

    # Row 4: 003 | Gamma | Raj | $60K
    row4_data = ["003", "Gamma", "Raj", "$60K"]
    row4 = table.rows[3]
    for col_idx, text in enumerate(row4_data):
        cell = row4.cells[col_idx]
        cell.paragraphs[0].clear()
        cell.paragraphs[0].add_run(text)

    # Row 5: 004 | Delta | Ana | $45K
    row5_data = ["004", "Delta", "Ana", "$45K"]
    row5 = table.rows[4]
    for col_idx, text in enumerate(row5_data):
        cell = row5.cells[col_idx]
        cell.paragraphs[0].clear()
        cell.paragraphs[0].add_run(text)

    # Add a footer paragraph
    doc.add_paragraph("")
    doc.add_paragraph("Note: All budget figures are approximate estimates pending final approval.")

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
