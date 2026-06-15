"""
Reward Script: Set third row background to light orange and make second column italic
Task ID: writer_tbl_048
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): All 4 cells in Row 3 have light orange (FFCC99) background fill
  Component 2 (0.4): All 5 cells in column 2 have italic text in their runs
  Component 3 (0.1): Cell (Row3, Col2) has BOTH light orange background AND italic (intersection check)
  Total: 1.0
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_048'
FILE_NAME = 'highlight_table.docx'

# Light orange color expected: #FFCC99 (or very close variants)
LIGHT_ORANGE_HEX = 'FFCC99'


def get_cell_bg_fill(cell):
    """
    Extract the background fill color (w:fill) from a table cell's shading element.
    Returns the hex color string (uppercase) or None if not set.
    """
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    if fill is None or fill.upper() in ('', 'NONE', 'AUTO'):
        return None
    return fill.upper()


def is_light_orange(hex_color):
    """
    Check if a hex color is light orange (FFCC99) with a small tolerance.
    Also accept known variants like FFD08C, FFCC88, etc. within a threshold.
    """
    if hex_color is None:
        return False
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        # Expected light orange: R=255, G=204, B=153
        # Allow tolerance of 20 per channel
        return abs(r - 255) <= 20 and abs(g - 204) <= 30 and abs(b - 153) <= 30
    except (ValueError, IndexError):
        return False


def cell_has_italic(cell):
    """
    Check if any run in the cell's paragraphs has italic=True.
    Returns True if at least one run with non-empty text is italic.
    """
    for para in cell.paragraphs:
        for run in para.runs:
            if run.text.strip() and run.font.italic is True:
                return True
    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Verify basic structure: 1 table with 5 rows and 4 columns
    try:
        if len(doc.tables) < 1:
            print("CRITICAL: No tables found in document")
            print("REWARD: 0.0")
            return 0.0
        table = doc.tables[0]
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        if num_rows < 5 or num_cols < 4:
            print(f"CRITICAL: Expected 5x4 table, found {num_rows}x{num_cols}")
            print("REWARD: 0.0")
            return 0.0
        print(f"INFO: Table found with {num_rows} rows and {num_cols} columns")
    except Exception as e:
        print(f"CRITICAL: Table structure check failed: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All 4 cells in Row 3 have light orange background (0.5 points)
    # Row 3 (index 2) should contain '002', 'Beta', 'Lee', '$75K' all with FFCC99 fill
    try:
        row3_orange_count = 0
        row3_cells_detail = []
        for col_idx in range(4):
            cell = table.cell(2, col_idx)
            bg = get_cell_bg_fill(cell)
            is_orange = is_light_orange(bg)
            row3_cells_detail.append(f"Col{col_idx+1}(text={cell.text.strip()!r}, fill={bg}, orange={is_orange})")
            if is_orange:
                row3_orange_count += 1

        if row3_orange_count == 4:
            print(f"PASS: Component 1 — All 4 Row 3 cells have light orange background: {row3_cells_detail} (0.5 pts)")
            total_score += 0.5
        elif row3_orange_count > 0:
            partial = round(0.5 * row3_orange_count / 4, 4)
            print(f"PARTIAL: Component 1 — {row3_orange_count}/4 Row 3 cells have light orange background: {row3_cells_detail} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No Row 3 cells have light orange background: {row3_cells_detail}")
    except Exception as e:
        print(f"ERROR: Component 1 — Row 3 background check failed: {e}")

    # Component 2: All 5 cells in column 2 (index 1) have italic text (0.4 points)
    # Column 2 cells: 'Project', 'Alpha', 'Beta', 'Gamma', 'Delta' — all should be italic
    try:
        col2_italic_count = 0
        col2_cells_detail = []
        for row_idx in range(5):
            cell = table.cell(row_idx, 1)
            has_it = cell_has_italic(cell)
            col2_cells_detail.append(f"Row{row_idx+1}(text={cell.text.strip()!r}, italic={has_it})")
            if has_it:
                col2_italic_count += 1

        if col2_italic_count == 5:
            print(f"PASS: Component 2 — All 5 column 2 cells have italic text: {col2_cells_detail} (0.4 pts)")
            total_score += 0.4
        elif col2_italic_count > 0:
            partial = round(0.4 * col2_italic_count / 5, 4)
            print(f"PARTIAL: Component 2 — {col2_italic_count}/5 column 2 cells have italic text: {col2_cells_detail} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No column 2 cells have italic text: {col2_cells_detail}")
    except Exception as e:
        print(f"ERROR: Component 2 — Column 2 italic check failed: {e}")

    # Component 3: Cell (Row3, Col2) has BOTH light orange background AND italic text (0.1 points)
    # This verifies the intersection: 'Beta' cell should have both properties
    try:
        cell_b3 = table.cell(2, 1)
        bg_b3 = get_cell_bg_fill(cell_b3)
        italic_b3 = cell_has_italic(cell_b3)
        b3_is_orange = is_light_orange(bg_b3)

        if b3_is_orange and italic_b3:
            print(f"PASS: Component 3 — Cell (Row3,Col2) text={cell_b3.text.strip()!r} has BOTH light orange bg ({bg_b3}) AND italic text (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 3 — Cell (Row3,Col2) text={cell_b3.text.strip()!r}: orange={b3_is_orange} (fill={bg_b3}), italic={italic_b3}")
    except Exception as e:
        print(f"ERROR: Component 3 — Intersection check failed: {e}")

    final_score = min(round(total_score, 4), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Main entry point: test canonical artifact path on the VM
file_path = f'{WORKDIR}/{FILE_NAME}'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
