"""
Initial Setup: Legal document with WHEREAS recital clauses (no hanging indents)
Task ID: writer_legal_038
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_038'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title ---
    title = doc.add_heading("MASTER SERVICES AGREEMENT", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Preamble ---
    preamble = doc.add_paragraph()
    preamble.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = preamble.add_run(
        "This Master Services Agreement (the \"Agreement\") is entered into as of "
        "March 15, 2025, by and between Meridian Technology Solutions, Inc., a Delaware "
        "corporation with its principal offices at 2400 Innovation Drive, Suite 800, "
        "San Francisco, CA 94105 (\"Provider\"), and Cascade Financial Group, LLC, a "
        "New York limited liability company with its principal offices at 550 Park "
        "Avenue, 12th Floor, New York, NY 10065 (\"Client\")."
    )
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    # --- Recitals heading ---
    recitals_heading = doc.add_heading("RECITALS", level=1)
    recitals_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- 5 WHEREAS clauses (NO hanging indent - plain paragraphs) ---
    whereas_clauses = [
        (
            "WHEREAS, Provider is engaged in the business of providing enterprise "
            "software development, cloud infrastructure management, and related "
            "technology consulting services to corporate clients throughout North "
            "America and Europe;"
        ),
        (
            "WHEREAS, Client operates a diversified financial services platform "
            "requiring comprehensive digital transformation of its legacy banking "
            "systems, customer relationship management tools, and regulatory "
            "compliance reporting infrastructure;"
        ),
        (
            "WHEREAS, Client has conducted an extensive vendor evaluation process "
            "spanning fourteen months and has determined that Provider possesses the "
            "technical expertise, industry certifications, and operational capacity "
            "necessary to deliver the services contemplated herein;"
        ),
        (
            "WHEREAS, the parties intend to establish a framework under which "
            "Provider shall deliver professional services, software licenses, and "
            "ongoing maintenance support pursuant to individual Statements of Work "
            "to be executed from time to time during the term of this Agreement;"
        ),
        (
            "WHEREAS, both parties acknowledge that the successful completion of "
            "the services described herein is contingent upon mutual cooperation, "
            "timely access to Client systems and personnel, and adherence to the "
            "project milestones and deliverable schedules set forth in each "
            "applicable Statement of Work."
        ),
    ]

    for clause_text in whereas_clauses:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        # No indentation at all - plain paragraph
        run = para.add_run(clause_text)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    # --- Agreement section ---
    now_therefore = doc.add_paragraph()
    now_therefore.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = now_therefore.add_run(
        "NOW, THEREFORE, in consideration of the mutual covenants and agreements "
        "set forth herein, and for other good and valuable consideration, the receipt "
        "and sufficiency of which are hereby acknowledged, the parties agree as follows:"
    )
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.bold = True

    # --- Article 1 ---
    art1 = doc.add_heading("ARTICLE 1 — DEFINITIONS", level=2)

    definitions = [
        ('"Confidential Information"', "means any and all non-public information, "
         "whether written, oral, electronic, or visual, disclosed by either party "
         "to the other in connection with this Agreement, including but not limited "
         "to trade secrets, business plans, financial data, customer lists, and "
         "proprietary technology."),
        ('"Deliverables"', "means the tangible and intangible work products, "
         "software code, documentation, reports, and other materials to be produced "
         "by Provider and delivered to Client as specified in each Statement of Work."),
        ('"Effective Date"', "means March 15, 2025, the date first written above."),
    ]

    for term, definition in definitions:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        run_term = para.add_run(term + " ")
        run_term.bold = True
        run_term.font.size = Pt(11)
        run_term.font.name = "Times New Roman"
        run_def = para.add_run(definition)
        run_def.font.size = Pt(11)
        run_def.font.name = "Times New Roman"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
