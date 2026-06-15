"""
Reward Script: Convert WHEREAS clauses to hanging indents
Task ID: writer_legal_038
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Each WHEREAS paragraph has left_indent = 1.5 inches (0.08 per clause)
  Component 2 (0.4): Each WHEREAS paragraph has first_line_indent = -1.5 inches (0.08 per clause)
  Component 3 (0.2): Non-WHEREAS paragraphs remain unaffected (no spurious indentation)
"""

import os

from docx import Document
from docx.shared import Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_038'

# Target values in EMU
TARGET_LEFT_INDENT = Inches(1.5)       # 1371600 EMU
TARGET_FIRST_LINE_INDENT = -Inches(1.5)  # -1371600 EMU
# Tolerance: ~0.05 inches in EMU for rounding
TOLERANCE = Inches(0.1)  # 91440 EMU


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Identify WHEREAS paragraphs (paragraphs starting with "WHEREAS")
    whereas_indices = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("WHEREAS"):
            whereas_indices.append(i)

    print(f"Found {len(whereas_indices)} WHEREAS paragraphs at indices: {whereas_indices}")

    if len(whereas_indices) == 0:
        print("FAIL: No WHEREAS paragraphs found in document")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: left_indent on WHEREAS paragraphs (0.4 points total)
    # Each WHEREAS clause scores 0.08 points (0.4 / 5)
    points_per_clause_c1 = 0.4 / max(len(whereas_indices), 1)
    try:
        c1_score = 0.0
        for idx in whereas_indices:
            para = doc.paragraphs[idx]
            pf = para.paragraph_format
            left_ind = pf.left_indent
            if left_ind is not None and abs(left_ind - TARGET_LEFT_INDENT) <= TOLERANCE:
                print(f"PASS: Para {idx} left_indent={left_ind} (expected ~{TARGET_LEFT_INDENT})")
                c1_score += points_per_clause_c1
            else:
                print(f"FAIL: Para {idx} left_indent={left_ind} (expected ~{TARGET_LEFT_INDENT})")
        print(f"Component 1 total: {c1_score:.2f}/0.40")
        if c1_score > 0:
            total_score += c1_score
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: first_line_indent on WHEREAS paragraphs (0.4 points total)
    # Each WHEREAS clause scores 0.08 points (0.4 / 5)
    points_per_clause_c2 = 0.4 / max(len(whereas_indices), 1)
    try:
        c2_score = 0.0
        for idx in whereas_indices:
            para = doc.paragraphs[idx]
            pf = para.paragraph_format
            first_ind = pf.first_line_indent
            if first_ind is not None and abs(first_ind - TARGET_FIRST_LINE_INDENT) <= TOLERANCE:
                print(f"PASS: Para {idx} first_line_indent={first_ind} (expected ~{TARGET_FIRST_LINE_INDENT})")
                c2_score += points_per_clause_c2
            else:
                print(f"FAIL: Para {idx} first_line_indent={first_ind} (expected ~{TARGET_FIRST_LINE_INDENT})")
        print(f"Component 2 total: {c2_score:.2f}/0.40")
        if c2_score > 0:
            total_score += c2_score
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All WHEREAS paragraphs have BOTH correct indents AND
    # non-WHEREAS paragraphs are unaffected (compound check, 0.2 points)
    # This only passes when the task is done correctly AND precisely.
    try:
        all_whereas_correct = (c1_score >= 0.39 and c2_score >= 0.39)
        non_whereas_ok = not False  # loop sentinel: set False if any non-WHEREAS has hanging indent
        for i, para in enumerate(doc.paragraphs):
            if i in whereas_indices:
                continue
            pf = para.paragraph_format
            left_ind = pf.left_indent
            first_ind = pf.first_line_indent
            # Check that non-WHEREAS paragraphs do NOT have the hanging indent pattern
            if (left_ind is not None and abs(left_ind - TARGET_LEFT_INDENT) <= TOLERANCE and
                    first_ind is not None and abs(first_ind - TARGET_FIRST_LINE_INDENT) <= TOLERANCE):
                print(f"FAIL: Para {i} (non-WHEREAS) unexpectedly has hanging indent: "
                      f"left={left_ind}, first={first_ind}")
                non_whereas_ok = False
        if all_whereas_correct and non_whereas_ok:
            print(f"PASS: Component 3 — all WHEREAS indented AND non-WHEREAS unaffected (0.20 pts)")
            total_score += 0.2
        elif not all_whereas_correct:
            print(f"FAIL: Component 3 — WHEREAS indentation incomplete, cannot award precision bonus")
        else:
            print(f"FAIL: Component 3 — some non-WHEREAS paragraphs have hanging indents")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
