"""
Initial Setup: Accept all tracked changes in thesis chapter document
Task ID: writer_rm_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

AUTHOR_PATEL = 'Dr. Patel'
AUTHOR_LIN = 'Student_Lin'
DATE_PATEL = '2025-11-10T14:30:00Z'
DATE_LIN = '2025-11-15T09:45:00Z'

rev_id = [100]

def next_id():
    rev_id[0] += 1
    return str(rev_id[0])


def make_run(parent, text, bold=False, italic=False):
    """Create a w:r element under parent."""
    r = etree.SubElement(parent, qn('w:r'))
    if bold or italic:
        rpr = etree.SubElement(r, qn('w:rPr'))
        if bold:
            etree.SubElement(rpr, qn('w:b'))
        if italic:
            etree.SubElement(rpr, qn('w:i'))
    t = etree.SubElement(r, qn('w:t'))
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    return r


def add_insertion(parent, text, author, date, bold=False, italic=False):
    """Add a tracked insertion (w:ins containing w:r) to parent element."""
    ins = etree.SubElement(parent, qn('w:ins'))
    ins.set(qn('w:id'), next_id())
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), date)
    r = etree.SubElement(ins, qn('w:r'))
    if bold or italic:
        rpr = etree.SubElement(r, qn('w:rPr'))
        if bold:
            etree.SubElement(rpr, qn('w:b'))
        if italic:
            etree.SubElement(rpr, qn('w:i'))
    t = etree.SubElement(r, qn('w:t'))
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    return ins


def add_deletion(parent, text, author, date, bold=False, italic=False):
    """Add a tracked deletion (w:del containing w:r with w:delText) to parent."""
    dl = etree.SubElement(parent, qn('w:del'))
    dl.set(qn('w:id'), next_id())
    dl.set(qn('w:author'), author)
    dl.set(qn('w:date'), date)
    r = etree.SubElement(dl, qn('w:r'))
    if bold or italic:
        rpr = etree.SubElement(r, qn('w:rPr'))
        if bold:
            etree.SubElement(rpr, qn('w:b'))
        if italic:
            etree.SubElement(rpr, qn('w:i'))
    dt = etree.SubElement(r, qn('w:delText'))
    dt.set(qn('xml:space'), 'preserve')
    dt.text = text
    return dl


def add_fmt_change_run(parent, text, new_bold=False, new_italic=False, author=AUTHOR_PATEL, date=DATE_PATEL):
    """Add a run with rPrChange (tracked formatting change) to parent."""
    r = etree.SubElement(parent, qn('w:r'))
    rpr = etree.SubElement(r, qn('w:rPr'))
    if new_bold:
        etree.SubElement(rpr, qn('w:b'))
    if new_italic:
        etree.SubElement(rpr, qn('w:i'))
    rpr_change = etree.SubElement(rpr, qn('w:rPrChange'))
    rpr_change.set(qn('w:id'), next_id())
    rpr_change.set(qn('w:author'), author)
    rpr_change.set(qn('w:date'), date)
    old_rpr = etree.SubElement(rpr_change, qn('w:rPr'))
    # old rPr is empty = no bold/italic before change
    t = etree.SubElement(r, qn('w:t'))
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    return r


def launch_gui(command, delay_sec=1.0):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    body = doc.element.body

    # ========== TITLE ==========
    title = doc.add_heading('Chapter 3: Research Methodology', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ========== Section 3.1 ==========
    doc.add_heading('3.1 Research Design', level=2)

    # --- Paragraph 1: CHANGE 1 (insertion by Dr. Patel) ---
    p1_elem = doc.add_paragraph()._element
    make_run(p1_elem,
        'This study employs a mixed-methods research design to investigate the impact of '
        'artificial intelligence integration on undergraduate learning outcomes in STEM disciplines. '
    )
    add_insertion(p1_elem,
        'specifically computer science and electrical engineering ',
        AUTHOR_PATEL, DATE_PATEL, italic=True
    )
    make_run(p1_elem,
        'across three participating universities during the 2024-2025 academic year.'
    )

    # --- Paragraph 2: CHANGE 2 (deletion by Student_Lin) ---
    p2_elem = doc.add_paragraph()._element
    make_run(p2_elem,
        'The research framework draws upon Creswell\'s convergent parallel design '
    )
    add_deletion(p2_elem,
        'which is widely recognized in educational research ',
        AUTHOR_LIN, DATE_LIN
    )
    make_run(p2_elem,
        'where quantitative and qualitative data are collected simultaneously and merged during interpretation.'
    )

    # ========== Section 3.2 ==========
    doc.add_heading('3.2 Participant Selection', level=2)

    # --- Paragraph 3: CHANGE 3 (insertion by Student_Lin) ---
    p3_elem = doc.add_paragraph()._element
    make_run(p3_elem,
        'A stratified random sampling technique was used to select 240 undergraduate students '
        'from three universities: Eastfield University, Westbrook Institute of Technology, and '
    )
    add_insertion(p3_elem,
        'Northern Plains State University',
        AUTHOR_LIN, DATE_LIN
    )
    make_run(p3_elem,
        '. Participants were distributed evenly across six course sections, with 40 students per section.'
    )

    # --- Paragraph 4: CHANGE 4 (deletion by Dr. Patel) ---
    p4_elem = doc.add_paragraph()._element
    make_run(p4_elem,
        'Inclusion criteria required participants to be enrolled full-time, '
    )
    add_deletion(p4_elem,
        'have some basic familiarity with computers, ',
        AUTHOR_PATEL, DATE_PATEL
    )
    make_run(p4_elem,
        'have no prior formal exposure to AI-assisted learning tools, and maintain a GPA of 2.5 or above.'
    )

    # ========== Section 3.3 ==========
    doc.add_heading('3.3 Data Collection Instruments', level=2)

    # --- Paragraph 5: CHANGE 5 (formatting change by Dr. Patel - bold) ---
    p5_elem = doc.add_paragraph()._element
    add_fmt_change_run(p5_elem,
        'Cognitive Load Assessment Scale (CLAS)',
        new_bold=True, author=AUTHOR_PATEL, date=DATE_PATEL
    )
    make_run(p5_elem,
        ' was administered at three time points during the semester: week 2 (baseline), '
        'week 8 (midpoint), and week 15 (endpoint). The instrument comprises 28 Likert-scale items '
        'measuring intrinsic, extraneous, and germane cognitive load dimensions.'
    )

    # --- Paragraph 6: CHANGE 6 (insertion by Student_Lin) ---
    p6_elem = doc.add_paragraph()._element
    make_run(p6_elem,
        'Semi-structured interviews were conducted with a purposive subsample of 36 participants '
        '(6 per section). Each interview lasted approximately 45 minutes and followed a protocol '
        'addressing four thematic areas: '
    )
    add_insertion(p6_elem,
        'perceived usefulness, ease of use, learning engagement, and self-efficacy',
        AUTHOR_LIN, DATE_LIN
    )
    make_run(p6_elem, '.')

    # ========== Section 3.4 ==========
    doc.add_heading('3.4 Experimental Procedure', level=2)

    # --- Paragraph 7: CHANGE 7 (deletion by Dr. Patel) ---
    p7_elem = doc.add_paragraph()._element
    make_run(p7_elem,
        'The experimental group received instruction supplemented by an AI tutoring system '
        '(TutorAI Pro v3.2) '
    )
    add_deletion(p7_elem,
        'developed in partnership with EduTech Solutions Inc. ',
        AUTHOR_PATEL, DATE_PATEL
    )
    make_run(p7_elem,
        'while the control group received traditional instruction with access to standard digital resources.'
    )

    # --- Paragraph 8: CHANGE 8 (formatting change by Student_Lin - italic) ---
    p8_elem = doc.add_paragraph()._element
    make_run(p8_elem, 'Intervention sessions were scheduled for ')
    add_fmt_change_run(p8_elem,
        'three 50-minute sessions per week',
        new_italic=True, author=AUTHOR_LIN, date=DATE_LIN
    )
    make_run(p8_elem,
        ' over a 12-week period, totaling 36 contact hours. Fidelity of implementation was '
        'monitored through weekly classroom observations and instructor logs.'
    )

    # ========== Section 3.5 ==========
    doc.add_heading('3.5 Data Analysis', level=2)

    # --- Paragraph 9: CHANGE 9 (insertion by Dr. Patel) ---
    p9_elem = doc.add_paragraph()._element
    make_run(p9_elem,
        'Quantitative data were analyzed using IBM SPSS Statistics (Version 29). '
    )
    add_insertion(p9_elem,
        'A two-way repeated measures ANOVA was employed to examine interaction effects between group and time. ',
        AUTHOR_PATEL, DATE_PATEL
    )
    make_run(p9_elem,
        'Effect sizes were calculated using partial eta-squared with thresholds of 0.01, 0.06, and 0.14 '
        'for small, medium, and large effects respectively.'
    )

    # --- Paragraph 10: CHANGE 10 (deletion by Student_Lin) ---
    p10_elem = doc.add_paragraph()._element
    make_run(p10_elem,
        'Qualitative data from interviews were transcribed verbatim and analyzed using '
        'thematic analysis following Braun and Clarke\'s six-phase framework. '
    )
    add_deletion(p10_elem,
        'The transcription was performed by a professional transcription service and verified by the research team. ',
        AUTHOR_LIN, DATE_LIN
    )
    make_run(p10_elem,
        'Two independent coders analyzed the data with an inter-rater reliability (Cohen\'s kappa) of 0.87.'
    )

    # ========== Section 3.6 ==========
    doc.add_heading('3.6 Ethical Considerations', level=2)

    # --- Paragraph 11: CHANGE 11 (insertion by Dr. Patel) ---
    p11_elem = doc.add_paragraph()._element
    make_run(p11_elem,
        'This study received approval from the Institutional Review Board (Protocol #IRB-2024-0847). '
    )
    add_insertion(p11_elem,
        'Written informed consent was obtained from all participants prior to data collection. ',
        AUTHOR_PATEL, DATE_PATEL
    )
    make_run(p11_elem,
        'Participants were informed of their right to withdraw at any time without academic penalty. '
        'All data were anonymized using a double-blind coding system.'
    )

    # --- Paragraph 12: CHANGE 12 (formatting change by Dr. Patel - bold) ---
    p12_elem = doc.add_paragraph()._element
    add_fmt_change_run(p12_elem,
        'Data retention policy',
        new_bold=True, author=AUTHOR_PATEL, date=DATE_PATEL
    )
    make_run(p12_elem,
        ' follows university guidelines requiring secure storage for a minimum of five years '
        'following publication, after which all identifiable records will be permanently destroyed.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
