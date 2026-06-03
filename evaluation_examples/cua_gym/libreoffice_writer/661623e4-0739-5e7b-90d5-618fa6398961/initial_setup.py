"""
Initial Setup: Insert the document title as a field in the header
Task ID: writer_tm_054
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_054'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_document_title(doc, title):
    """Set the Title document property in core properties."""
    doc.core_properties.title = title


def create_initial():
    doc = Document()

    # --- Set Document Properties ---
    set_document_title(doc, "Cloud Migration Proposal")
    doc.core_properties.author = "Sarah Chen"
    doc.core_properties.subject = "Cloud Infrastructure Migration"

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Header: MUST be empty (task is to insert title field here) ---
    header = section.header
    header.is_linked_to_previous = False
    # Ensure header exists but is empty
    if header.paragraphs:
        header.paragraphs[0].text = ""

    # --- Footer with page numbers ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fp.add_run("Page ")
    r1 = fp.add_run()
    r1._element.append(r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}))
    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    r3._element.append(r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}))

    # --- Document Body ---
    # Title heading
    heading = doc.add_heading("Cloud Migration Proposal", level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This proposal outlines a comprehensive strategy for migrating our on-premises "
        "infrastructure to Amazon Web Services (AWS) and Microsoft Azure cloud platforms. "
        "The migration will be executed in three phases over 18 months, targeting a 35% "
        "reduction in total infrastructure costs while improving system reliability and scalability."
    )

    # Background
    doc.add_heading("Background", level=1)
    doc.add_paragraph(
        "Meridian Technologies currently operates 142 physical servers across two data centers "
        "in Portland, Oregon and Richmond, Virginia. Our legacy infrastructure, originally "
        "provisioned in 2017, faces increasing maintenance costs and diminishing vendor support. "
        "Key applications including the CRM platform, financial reporting suite, and customer-facing "
        "web services experience periodic outages averaging 4.2 hours per quarter."
    )
    doc.add_paragraph(
        "The board of directors approved an initial feasibility study in Q3 2025, which confirmed "
        "that cloud migration would yield significant operational benefits. This proposal details "
        "the implementation plan, resource requirements, and expected return on investment."
    )

    # Scope
    doc.add_heading("Scope of Migration", level=1)
    doc.add_paragraph(
        "The migration encompasses the following workloads and services:"
    )
    items = [
        "Customer Relationship Management (CRM) - Salesforce integration layer and custom modules",
        "Financial Reporting Suite - 12 reporting databases and ETL pipelines",
        "Public-Facing Web Applications - 8 web services handling 2.3M daily requests",
        "Internal Collaboration Tools - Email servers, document management, and intranet",
        "Development and Testing Environments - CI/CD pipelines and staging infrastructure",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    # Timeline
    doc.add_heading("Implementation Timeline", level=1)

    # Phase 1
    doc.add_heading("Phase 1: Assessment and Planning (Months 1-4)", level=2)
    doc.add_paragraph(
        "Conduct detailed workload analysis, dependency mapping, and security assessment. "
        "Establish cloud landing zones in AWS (us-west-2) and Azure (West US 2). "
        "Estimated cost: $185,000 including external consulting fees."
    )

    # Phase 2
    doc.add_heading("Phase 2: Migration Execution (Months 5-14)", level=2)
    doc.add_paragraph(
        "Execute staged migration using a lift-and-shift approach for legacy workloads and "
        "re-platforming for cloud-native candidates. Priority order: development environments, "
        "internal tools, web applications, financial systems, and CRM platform. "
        "Estimated cost: $620,000 including temporary dual-running expenses."
    )

    # Phase 3
    doc.add_heading("Phase 3: Optimization (Months 15-18)", level=2)
    doc.add_paragraph(
        "Decommission on-premises hardware, optimize cloud resource allocation using reserved "
        "instances and auto-scaling policies, and implement comprehensive monitoring with "
        "CloudWatch and Azure Monitor dashboards. Estimated cost: $95,000."
    )

    # Budget table
    doc.add_heading("Budget Summary", level=1)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers_row = ["Category", "Estimated Cost", "Notes"]
    for i, h in enumerate(headers_row):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    budget_data = [
        ["Cloud Platform Licenses", "$340,000/year", "AWS + Azure combined"],
        ["Migration Services", "$620,000", "One-time consulting and labor"],
        ["Assessment & Planning", "$185,000", "Phase 1 external resources"],
        ["Post-Migration Optimization", "$95,000", "Months 15-18"],
        ["Contingency (15%)", "$186,000", "Risk buffer"],
    ]
    for r, row_data in enumerate(budget_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # Risk Assessment
    doc.add_heading("Risk Assessment", level=1)
    doc.add_paragraph(
        "Key risks identified during the feasibility study include data loss during migration "
        "(mitigated by parallel running and incremental sync), extended downtime windows "
        "(mitigated by blue-green deployment strategy), and staff skill gaps (mitigated by "
        "AWS and Azure certification training program for 24 IT staff members)."
    )

    # Conclusion
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "The proposed cloud migration represents a strategic investment that will modernize "
        "Meridian Technologies' IT infrastructure, reduce operational costs by an estimated "
        "$480,000 annually after Year 2, and position the organization for future growth. "
        "We recommend immediate approval to begin Phase 1 activities."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
