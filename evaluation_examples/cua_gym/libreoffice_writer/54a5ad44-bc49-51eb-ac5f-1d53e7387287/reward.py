"""
Reward Script: Insert section break with two-column layout for parallel translations
Task ID: writer_acad_073
Domain: libreoffice_writer
Scoring:
  Component 1: Document has multiple sections (0.2 pts)
  Component 2: At least one section has two-column layout (0.3 pts)
  Component 3: Two-column section uses continuous section break (0.2 pts)
  Component 4: Non-two-column sections remain single-column (0.15 pts)
  Component 5: Two-column section contains parallel translation content (0.15 pts)
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_073'


def persist_app_state(domain):
    """Save any unsaved changes in LibreOffice Writer."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Cannot import python-docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    sections = doc.sections
    num_sections = len(sections)

    # Component 1: Document has multiple sections (0.2 points)
    # Initial env has only 1 section. Golden env should have 3 (before, two-col, after).
    try:
        if num_sections >= 3:
            print(f"PASS: Component 1 -- Document has {num_sections} sections (>=3) (0.2 pts)")
            total_score += 0.2
        elif num_sections == 2:
            print(f"PARTIAL: Component 1 -- Document has 2 sections (expected >=3) (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 1 -- Document has only {num_sections} section(s), expected >=3")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: At least one section has two-column layout (0.3 points)
    # This is the core task requirement. Check w:cols w:num="2" in section properties.
    try:
        two_col_section_indices = []
        for i, s in enumerate(sections):
            sect_el = s._sectPr
            cols_el = sect_el.find(qn('w:cols'))
            if cols_el is not None:
                num_attr = cols_el.get(qn('w:num'))
                if num_attr is not None and int(num_attr) == 2:
                    two_col_section_indices.append(i)

        if len(two_col_section_indices) > 0:
            print(f"PASS: Component 2 -- Found two-column section(s) at index(es): {two_col_section_indices} (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 -- No section with two-column layout found")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: The two-column section uses continuous section break (0.2 points)
    # A proper section break for column change within a page should be CONTINUOUS type.
    try:
        if len(two_col_section_indices) > 0:
            continuous_count = 0
            for idx in two_col_section_indices:
                s = sections[idx]
                sect_el = s._sectPr
                type_el = sect_el.find(qn('w:type'))
                if type_el is not None:
                    val = type_el.get(qn('w:val'))
                    if val == 'continuous':
                        continuous_count += 1
                        print(f"  Section {idx}: type=continuous (correct)")
                    else:
                        print(f"  Section {idx}: type={val} (expected continuous)")
                else:
                    # Default section break type (NEW_PAGE) -- not ideal for column change
                    print(f"  Section {idx}: type=default/newPage (expected continuous)")

            if continuous_count == len(two_col_section_indices):
                print(f"PASS: Component 3 -- All two-column sections use continuous break (0.2 pts)")
                total_score += 0.2
            elif continuous_count > 0:
                print(f"PARTIAL: Component 3 -- {continuous_count}/{len(two_col_section_indices)} two-col sections are continuous (0.1 pts)")
                total_score += 0.1
            else:
                print(f"FAIL: Component 3 -- No two-column section uses continuous break")
        else:
            print(f"FAIL: Component 3 -- No two-column section exists, cannot check break type")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Sections before and after two-column section are single-column (0.15 points)
    # The rest of the document should remain single-column.
    try:
        if len(two_col_section_indices) > 0 and num_sections >= 3:
            multi_col_other_sections = []
            for i, s in enumerate(sections):
                if i in two_col_section_indices:
                    continue
                sect_el = s._sectPr
                cols_el = sect_el.find(qn('w:cols'))
                if cols_el is not None:
                    num_attr = cols_el.get(qn('w:num'))
                    if num_attr is not None and int(num_attr) > 1:
                        multi_col_other_sections.append(i)
                        print(f"  Section {i}: has {num_attr} columns (expected 1)")

            if len(multi_col_other_sections) == 0:
                print(f"PASS: Component 4 -- All non-two-column sections are single-column (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 4 -- Some non-two-column sections have multiple columns")
        else:
            print(f"FAIL: Component 4 -- Insufficient sections to verify column layout context")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    # Component 5: Two-column section contains parallel translation content (0.15 points)
    # The two-column section should contain Arabic and Latin translation passages.
    # We check by finding paragraphs within the two-column section and looking for
    # translation-related keywords.
    try:
        if len(two_col_section_indices) > 0:
            # Find paragraphs belonging to the two-column section.
            # In docx, inline section breaks (in paragraph properties) define the END
            # of a section. We need to find which paragraphs are between the section breaks.
            body = doc.element.body
            para_elements = body.findall(qn('w:p'))

            # Build section boundaries by finding paragraphs with sectPr
            section_break_para_indices = []
            for pi, p_el in enumerate(para_elements):
                pPr = p_el.find(qn('w:pPr'))
                if pPr is not None:
                    sect_in_pPr = pPr.find(qn('w:sectPr'))
                    if sect_in_pPr is not None:
                        # Check if this section is the two-column one
                        cols_el = sect_in_pPr.find(qn('w:cols'))
                        if cols_el is not None:
                            num_attr = cols_el.get(qn('w:num'))
                            if num_attr is not None and int(num_attr) == 2:
                                section_break_para_indices.append(pi)

            # The two-column sectPr in a paragraph means all paragraphs from the
            # previous section break up to (and including) this one are in that section.
            # Find the previous section break to determine the start of the two-col range.
            all_break_indices = []
            for pi, p_el in enumerate(para_elements):
                pPr = p_el.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    all_break_indices.append(pi)

            has_translation_content = False
            if section_break_para_indices:
                two_col_break_idx = section_break_para_indices[0]
                # Find the break before this one
                prev_break_idx = -1
                for bi in all_break_indices:
                    if bi < two_col_break_idx:
                        prev_break_idx = bi
                    else:
                        break

                # Paragraphs in the two-column section: from prev_break_idx+1 to two_col_break_idx
                start_para = prev_break_idx + 1
                end_para = two_col_break_idx

                two_col_text = ""
                for pi in range(start_para, end_para + 1):
                    if pi < len(doc.paragraphs):
                        two_col_text += doc.paragraphs[pi].text + " "

                two_col_text_lower = two_col_text.lower()
                # Check for translation-related content markers
                has_arabic = "arabic" in two_col_text_lower or any(
                    '\u0600' <= c <= '\u06FF' for c in two_col_text
                )
                has_latin = "latin" in two_col_text_lower or "substantia" in two_col_text_lower or "subiecto" in two_col_text_lower

                if has_arabic and has_latin:
                    print(f"PASS: Component 5 -- Two-column section contains both Arabic and Latin translation content (0.15 pts)")
                    total_score += 0.15
                elif has_arabic or has_latin:
                    print(f"PARTIAL: Component 5 -- Two-column section contains some translation content (0.1 pts)")
                    total_score += 0.1
                else:
                    print(f"FAIL: Component 5 -- Two-column section does not contain translation content. Text excerpt: {two_col_text[:200]}")
            else:
                print(f"FAIL: Component 5 -- Could not locate two-column section break in paragraph elements")
        else:
            print(f"FAIL: Component 5 -- No two-column section exists")
    except Exception as e:
        print(f"ERROR: Component 5 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
