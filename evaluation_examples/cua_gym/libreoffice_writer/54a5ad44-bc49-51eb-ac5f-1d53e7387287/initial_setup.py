"""
Initial Setup: Insert a section break and set two-column layout for parallel translations
Task ID: writer_acad_073
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_073'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title Page ---
    for _ in range(4):
        doc.add_paragraph('')

    title = doc.add_heading('', level=0)
    run = title.add_run('Cross-Cultural Hermeneutics:\nA Comparative Study of Translation Methodologies\nin Medieval Arabic-Latin Philosophical Texts')
    run.font.size = Pt(26)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = author.add_run('A Dissertation Submitted in Partial Fulfillment\n'
                       'of the Requirements for the Degree of\nDoctor of Philosophy')
    r.font.size = Pt(12)

    doc.add_paragraph('')

    author2 = doc.add_paragraph()
    author2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r2 = author2.add_run('Elena Vasquez-Rahman\nDepartment of Comparative Literature\nUniversity of Cambridge\n2025')
    r2.font.size = Pt(12)

    doc.add_page_break()

    # --- Abstract ---
    abs_heading = doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This dissertation examines the transmission of Aristotelian philosophical concepts '
        'through the Arabic-Latin translation movement of the 12th and 13th centuries. By '
        'comparing parallel translations of key passages from Ibn Rushd\'s Tahafut al-Tahafut '
        'and its Latin rendering by Calo Calonymos (1328), this study reveals how translators '
        'navigated the conceptual gaps between Greek metaphysics, Arabic falsafa, and Latin '
        'scholastic terminology. The analysis focuses on three pivotal concepts\u2014substance '
        '(ousia/jawhar/substantia), causation (aitia/sabab/causa), and potentiality '
        '(dynamis/quwwa/potentia)\u2014demonstrating that translation choices were not merely '
        'linguistic but carried profound philosophical implications that shaped the reception '
        'of Aristotelian thought in medieval Europe.'
    )

    doc.add_page_break()

    # --- Chapter 1: Introduction ---
    doc.add_heading('Chapter 1: Introduction', level=1)
    doc.add_heading('1.1 Background and Motivation', level=2)
    doc.add_paragraph(
        'The medieval Arabic-Latin translation movement represents one of the most significant '
        'episodes of cross-cultural intellectual exchange in human history. Beginning in Toledo '
        'and extending through centers of learning in Sicily, Provence, and the Italian '
        'city-states, this movement transformed the intellectual landscape of Western Europe by '
        'making available the philosophical and scientific heritage of classical antiquity, as '
        'mediated and enriched by centuries of Arabic scholarship.'
    )
    doc.add_paragraph(
        'While considerable scholarly attention has been devoted to the historical and '
        'institutional aspects of this translation movement (Burnett, 2001; d\'Alverny, 1982; '
        'Hasse, 2006), comparatively less work has examined the micro-level linguistic and '
        'conceptual transformations that occurred in the translation process itself. How did '
        'translators handle terms for which no direct equivalent existed in the target language? '
        'What philosophical assumptions guided their choices? And how did these choices, in turn, '
        'shape the understanding of the translated texts in their new intellectual context?'
    )

    doc.add_heading('1.2 Research Questions', level=2)
    doc.add_paragraph(
        'This dissertation addresses three interconnected research questions:'
    )
    doc.add_paragraph(
        'First, how did the translation of key Aristotelian metaphysical concepts from Arabic '
        'to Latin reflect differing philosophical frameworks?',
        style='List Number'
    )
    doc.add_paragraph(
        'Second, what systematic patterns can be identified in the translation strategies '
        'employed by Calo Calonymos in his Latin rendering of Ibn Rushd\'s Tahafut al-Tahafut?',
        style='List Number'
    )
    doc.add_paragraph(
        'Third, how did these translation choices influence the reception and interpretation '
        'of Aristotelian philosophy in the Latin West?',
        style='List Number'
    )

    doc.add_heading('1.3 Methodology', level=2)
    doc.add_paragraph(
        'The methodology employed in this study combines close textual analysis with broader '
        'philosophical contextualization. Following the comparative philological approach '
        'pioneered by Endress (1966) and refined by Gutas (1998), we examine parallel passages '
        'in the Arabic original and Latin translation, identifying patterns of semantic shift, '
        'conceptual adaptation, and interpretive interpolation.'
    )

    doc.add_page_break()

    # --- Chapter 2: Literature Review ---
    doc.add_heading('Chapter 2: Literature Review', level=1)
    doc.add_heading('2.1 The Translation Movement: Historical Overview', level=2)
    doc.add_paragraph(
        'The Arabic-Latin translation movement can be divided into three broad phases. The '
        'first phase (c. 1130\u20131187), centered in Toledo under the patronage of Archbishop '
        'Raymond, saw the translation of major astronomical and mathematical works. The second '
        'phase (c. 1187\u20131250) expanded to include philosophical and medical texts, with '
        'translators such as Michael Scot and Hermann of Carinthia producing Latin versions of '
        'Avicenna\'s and Averroes\'s major works. The third phase (c. 1250\u20131350) was '
        'characterized by increased attention to textual accuracy and philosophical nuance, '
        'exemplified by William of Moerbeke\'s direct translations from Greek and Calo '
        'Calonymos\'s careful rendering of Ibn Rushd.'
    )

    doc.add_heading('2.2 Translation Theory and Medieval Practice', level=2)
    doc.add_paragraph(
        'Modern translation theory, from Nida\'s (1964) distinction between formal and dynamic '
        'equivalence to Venuti\'s (1995) concepts of domestication and foreignization, provides '
        'useful analytical frameworks for understanding medieval translation practice. However, '
        'as Burnett (2006) has cautioned, we must be careful not to impose anachronistic '
        'categories on medieval translators whose understanding of their task differed '
        'fundamentally from modern conceptions of translation.'
    )

    doc.add_heading('2.3 Philosophical Terminology across Traditions', level=2)
    doc.add_paragraph(
        'The transmission of philosophical terminology across linguistic and cultural boundaries '
        'has been the subject of extensive scholarship. Frank (1956) and Wolfson (1976) mapped '
        'the Greek-Arabic terminological correspondences, while Daiber (1980) and Endress (1966) '
        'traced the further transmission into Latin. More recently, Adamson (2015) and Taylor '
        '(2012) have examined specific case studies of conceptual transformation in the '
        'translation process.'
    )

    doc.add_page_break()

    # --- Chapter 3: Theoretical Framework ---
    doc.add_heading('Chapter 3: Theoretical Framework', level=1)
    doc.add_heading('3.1 Conceptual Mapping Theory', level=2)
    doc.add_paragraph(
        'This study employs a theoretical framework that combines insights from conceptual '
        'metaphor theory (Lakoff & Johnson, 1980) with the notion of semantic fields as applied '
        'to philosophical terminology. We propose the concept of "philosophical translation '
        'space"\u2014a theoretical construct that maps the semantic and conceptual relationships '
        'between terms across languages and philosophical traditions.'
    )

    doc.add_heading('3.2 Layers of Mediation', level=2)
    doc.add_paragraph(
        'A crucial feature of the Arabic-Latin transmission of Aristotelian philosophy is its '
        'multi-layered character. The original Greek text passed through at least two stages of '
        'translation and interpretation before reaching Latin readers: first from Greek to '
        'Arabic (often via Syriac), and then from Arabic to Latin. Each stage introduced its '
        'own layer of interpretation, shaped by the philosophical assumptions and linguistic '
        'resources of the translator and the intellectual community for which the translation '
        'was intended.'
    )

    doc.add_page_break()

    # --- Chapter 4: Comparative Analysis ---
    doc.add_heading('Chapter 4: Comparative Analysis of Key Passages', level=1)
    doc.add_heading('4.1 Methodological Approach', level=2)
    doc.add_paragraph(
        'In this chapter, we present the core analytical findings of our study through a '
        'detailed comparison of parallel passages from Ibn Rushd\'s Tahafut al-Tahafut and '
        'Calo Calonymos\'s Latin translation (Destructio Destructionum, 1328). The passages '
        'have been selected for their philosophical significance and the richness of the '
        'translation challenges they present.'
    )

    doc.add_heading('4.2 Passage Analysis: On Substance (jawhar/substantia)', level=2)
    doc.add_paragraph(
        'The concept of substance provides our first case study. In the following passage, '
        'Ibn Rushd discusses the relationship between substance and existence, a topic central '
        'to the dispute between Avicennan and Aristotelian metaphysics. The Arabic text and '
        'its Latin translation are presented below for comparison.'
    )

    # This is the section where the user needs to ADD a two-column layout
    # In the INITIAL state, everything is single-column
    doc.add_paragraph(
        'Arabic (Ibn Rushd, Tahafut al-Tahafut, ed. Bouyges, p. 302):'
    )
    doc.add_paragraph(
        '\u0648\u0623\u0645\u0627 \u0642\u0648\u0644\u0647 \u0625\u0646 \u0627\u0644\u062c\u0648\u0647\u0631 '
        '\u0647\u0648 \u0627\u0644\u0645\u0648\u062c\u0648\u062f \u0627\u0644\u0630\u064a \u0644\u0627 '
        '\u064a\u0643\u0648\u0646 \u0641\u064a \u0645\u0648\u0636\u0648\u0639\u060c \u0641\u0647\u0630\u0627 '
        '\u0642\u0648\u0644 \u0635\u062d\u064a\u062d \u0625\u0646 \u0623\u0631\u064a\u062f \u0628\u0647 '
        '\u0627\u0644\u062c\u0648\u0647\u0631 \u0627\u0644\u0623\u0648\u0644\u060c \u0648\u0647\u0648 '
        '\u0627\u0644\u0641\u0631\u062f \u0627\u0644\u0645\u0634\u0627\u0631 \u0625\u0644\u064a\u0647. '
        '\u0648\u0623\u0645\u0627 \u0627\u0644\u062c\u0648\u0647\u0631 \u0627\u0644\u062b\u0627\u0646\u064a\u060c '
        '\u0648\u0647\u0648 \u0627\u0644\u0646\u0648\u0639 \u0648\u0627\u0644\u062c\u0646\u0633\u060c '
        '\u0641\u0647\u0648 \u0645\u0648\u062c\u0648\u062f \u0641\u064a \u0627\u0644\u0623\u0641\u0631\u0627\u062f '
        '\u0648\u0644\u064a\u0633 \u0641\u064a \u0645\u0648\u0636\u0648\u0639.'
    )
    doc.add_paragraph(
        'Latin (Calo Calonymos, Destructio Destructionum, Venice 1497, f. 87r):'
    )
    doc.add_paragraph(
        'Quod autem dicit quod substantia est ens quod non est in subiecto, hoc est dictum '
        'verum si intelligatur per substantiam substantia prima, quae est individuum '
        'demonstratum. Substantia autem secunda, quae est species et genus, est existens in '
        'individuis et non in subiecto.'
    )

    doc.add_heading('4.3 Passage Analysis: On Causation (sabab/causa)', level=2)
    doc.add_paragraph(
        'Our second case study focuses on the translation of causal terminology. The concept '
        'of sabab in Arabic philosophical usage encompasses a broader semantic range than the '
        'Latin causa, creating significant challenges for the translator. The following parallel '
        'passages illustrate the strategies employed by Calonymos in navigating these differences.'
    )
    doc.add_paragraph(
        'In this passage, Ibn Rushd critiques al-Ghazali\'s understanding of efficient '
        'causation. The Arabic term al-sabab al-fa\u02bbil (the efficient cause) is rendered '
        'by Calonymos as causa efficiens, following the established Scholastic terminology '
        'derived from Aristotle\'s Physics II.3.'
    )

    doc.add_heading('4.4 Passage Analysis: On Potentiality (quwwa/potentia)', level=2)
    doc.add_paragraph(
        'The third case study examines the translation of the Aristotelian concept of '
        'potentiality. The Arabic term quwwa, derived from the root meaning "strength" or '
        '"power," carries connotations absent from the Greek dynamis and the Latin potentia. '
        'This semantic surplus created opportunities for philosophical innovation that were '
        'sometimes lost in the Latin translation.'
    )

    doc.add_page_break()

    # --- Chapter 5: Conclusions ---
    doc.add_heading('Chapter 5: Conclusions', level=1)
    doc.add_heading('5.1 Summary of Findings', level=2)
    doc.add_paragraph(
        'This dissertation has demonstrated that the Arabic-Latin translation of Aristotelian '
        'philosophical concepts was far more than a mechanical transfer of meaning between '
        'languages. Through detailed analysis of parallel passages from Ibn Rushd\'s Tahafut '
        'al-Tahafut and its Latin translation by Calo Calonymos, we have shown that translation '
        'choices were shaped by, and in turn shaped, the philosophical frameworks within which '
        'the translators operated.'
    )

    doc.add_heading('5.2 Implications for Translation Studies', level=2)
    doc.add_paragraph(
        'The findings of this study have significant implications for our understanding of '
        'philosophical translation as an intellectual practice. First, they confirm the '
        'impossibility of "neutral" translation in philosophy: every rendering of a '
        'philosophical term is simultaneously an interpretation. Second, they demonstrate the '
        'productive role of translation in philosophical innovation, as translators created new '
        'conceptual possibilities through their terminological choices.'
    )

    doc.add_heading('5.3 Directions for Future Research', level=2)
    doc.add_paragraph(
        'Several avenues for future research emerge from this study. First, the comparative '
        'methodology developed here could be applied to other translator-philosopher pairs, '
        'such as Michael Scot\'s renderings of Avicenna or Hermann of Carinthia\'s translations '
        'of Abu Ma\'shar. Second, the concept of "philosophical translation space" could be '
        'extended to encompass other major episodes of cross-cultural philosophical transmission, '
        'such as the Sanskrit-Chinese transmission of Buddhist philosophy or the modern '
        'East-West philosophical dialogue.'
    )

    # --- Bibliography ---
    doc.add_page_break()
    doc.add_heading('Bibliography', level=1)
    refs = [
        'Adamson, P. (2015). "Arabic into Latin: The Reception of Arabic Philosophy into '
        'Western Europe." In The Cambridge Companion to Arabic Philosophy, ed. P. Adamson and '
        'R. Taylor, pp. 370\u2013396. Cambridge: Cambridge University Press.',
        'Burnett, C. (2001). "The Coherence of the Arabic-Latin Translation Program in Toledo '
        'in the Twelfth Century." Science in Context 14(1\u20132): 249\u2013288.',
        'Burnett, C. (2006). "Arabic-Latin Translation Program." In Medieval Science, Technology, '
        'and Medicine: An Encyclopedia, ed. T. Glick et al., pp. 40\u201344.',
        'Daiber, H. (1980). Aetius Arabus: Die Vorsokratiker in arabischer \u00dcberlieferung. '
        'Wiesbaden: Steiner.',
        'd\'Alverny, M.-T. (1982). "Translations and Translators." In Renaissance and Renewal '
        'in the Twelfth Century, ed. R. Benson and G. Constable, pp. 421\u2013462.',
        'Endress, G. (1966). Die arabischen \u00dcbersetzungen von Aristoteles\' Schrift De Caelo. '
        'Frankfurt: Klostermann.',
        'Frank, R. (1956). "The Origin of the Arabic Philosophical Term anniyya." '
        'Cahiers de Byrsa 6: 181\u2013201.',
        'Gutas, D. (1998). Greek Thought, Arabic Culture: The Graeco-Arabic Translation Movement '
        'in Baghdad. London: Routledge.',
        'Hasse, D.N. (2006). "The Social Conditions of the Arabic-Hebrew-Latin Translation '
        'Movements." In Wissen \u00fcber Grenzen, ed. A. Speer and L. Wegener, pp. 68\u201386.',
        'Wolfson, H.A. (1976). The Philosophy of the Kalam. Cambridge, MA: Harvard University Press.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
