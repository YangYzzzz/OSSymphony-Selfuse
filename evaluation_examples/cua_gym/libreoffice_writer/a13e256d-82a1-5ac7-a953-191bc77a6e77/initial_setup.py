"""
Initial Setup: Two numbered lists with second list restarting at 1
Task ID: writer_lec_006
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_numbered_list_item(doc, text, num_id, restart=False):
    """Add a numbered list item with explicit numId. If restart=True, restarts at 1."""
    para = doc.add_paragraph()
    # Set the paragraph style
    para.style = doc.styles['List Number']

    # Clear any auto-generated runs and set text
    para.clear()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # Set numPr explicitly
    pPr = para._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numPr.append(ilvl)
    numId_elem = OxmlElement('w:numId')
    numId_elem.set(qn('w:val'), str(num_id))
    numPr.append(numId_elem)
    pPr.append(numPr)

    return para


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add a title
    heading = doc.add_heading('Quarterly Marketing Strategy Review', level=1)
    for run in heading.runs:
        run.font.size = Pt(16)

    # Intro paragraph
    intro = doc.add_paragraph(
        'The following action items were identified during our Q2 planning session '
        'held on March 28, 2025. Each department lead is responsible for tracking '
        'progress on their assigned items.'
    )
    intro.paragraph_format.space_after = Pt(6)

    # Sub-heading for first list
    sub1 = doc.add_heading('Priority Action Items', level=2)

    # We need to set up numbering in the document XML.
    # Access the numbering part and create two abstract numbering definitions.
    # One for the first list, and a separate one for the second list (restarting).

    numbering_part = doc.part.numbering_part
    numbering_elem = numbering_part._element

    # Create abstract numbering definition for list 1
    abstract_num1 = OxmlElement('w:abstractNum')
    abstract_num1.set(qn('w:abstractNumId'), '10')
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')
    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    lvl.append(start)
    numFmt = OxmlElement('w:numFmt')
    numFmt.set(qn('w:val'), 'decimal')
    lvl.append(numFmt)
    lvlText = OxmlElement('w:lvlText')
    lvlText.set(qn('w:val'), '%1.')
    lvl.append(lvlText)
    lvlJc = OxmlElement('w:lvlJc')
    lvlJc.set(qn('w:val'), 'left')
    lvl.append(lvlJc)
    abstract_num1.append(lvl)

    # Create abstract numbering definition for list 2 (restarting at 1)
    abstract_num2 = OxmlElement('w:abstractNum')
    abstract_num2.set(qn('w:abstractNumId'), '11')
    lvl2 = OxmlElement('w:lvl')
    lvl2.set(qn('w:ilvl'), '0')
    start2 = OxmlElement('w:start')
    start2.set(qn('w:val'), '1')
    lvl2.append(start2)
    numFmt2 = OxmlElement('w:numFmt')
    numFmt2.set(qn('w:val'), 'decimal')
    lvl2.append(numFmt2)
    lvlText2 = OxmlElement('w:lvlText')
    lvlText2.set(qn('w:val'), '%1.')
    lvl2.append(lvlText2)
    lvlJc2 = OxmlElement('w:lvlJc')
    lvlJc2.set(qn('w:val'), 'left')
    lvl2.append(lvlJc2)
    abstract_num2.append(lvl2)

    # Insert abstract nums before any existing ones
    numbering_elem.insert(0, abstract_num1)
    numbering_elem.insert(1, abstract_num2)

    # Create num elements referencing the abstract nums
    num1 = OxmlElement('w:num')
    num1.set(qn('w:numId'), '20')
    abstractNumId1 = OxmlElement('w:abstractNumId')
    abstractNumId1.set(qn('w:val'), '10')
    num1.append(abstractNumId1)
    numbering_elem.append(num1)

    num2 = OxmlElement('w:num')
    num2.set(qn('w:numId'), '21')
    abstractNumId2 = OxmlElement('w:abstractNumId')
    abstractNumId2.set(qn('w:val'), '11')
    num2.append(abstractNumId2)
    numbering_elem.append(num2)

    # First numbered list (5 items) using numId 20
    list1_items = [
        'Finalize the social media campaign budget for Q3 with revised ROI projections',
        'Schedule follow-up meetings with the Apex Digital and BrightPath Agency partners',
        'Complete the competitive analysis report for the Southeast Asian market expansion',
        'Review and approve the updated brand guidelines document by April 15',
        'Coordinate with the analytics team to set up conversion tracking for the new landing pages',
    ]
    for item_text in list1_items:
        add_numbered_list_item(doc, item_text, num_id=20)

    # Regular paragraph between lists
    separator = doc.add_paragraph(
        'In addition to the priority items above, several supplementary tasks were '
        'discussed during the breakout sessions. These items should be addressed once '
        'the primary deliverables are underway.'
    )
    separator.paragraph_format.space_before = Pt(12)
    separator.paragraph_format.space_after = Pt(6)

    # Sub-heading for second list
    sub2 = doc.add_heading('Supplementary Tasks', level=2)

    # Second numbered list (3 items) using numId 21 - restarts at 1
    list2_items = [
        'Draft a proposal for the employee advocacy program pilot launch in May',
        'Update the content calendar to include the new product launch timeline',
        'Prepare a summary of customer feedback trends from the Q1 satisfaction survey',
    ]
    for item_text in list2_items:
        add_numbered_list_item(doc, item_text, num_id=21)

    # Closing paragraph
    closing = doc.add_paragraph(
        'All items should be tracked in the shared project management dashboard. '
        'Status updates are due every Friday by 4:00 PM EST.'
    )
    closing.paragraph_format.space_before = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
