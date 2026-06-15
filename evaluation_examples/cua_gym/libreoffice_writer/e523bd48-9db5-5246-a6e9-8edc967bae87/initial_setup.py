"""
Initial Setup: Policy manual document for task 'Add paragraph at end'
Task ID: writer_edit_051
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_051'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/policy_manual.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set document default styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Page 1: Title and Introduction ---
    title = doc.add_heading('Employee Policy Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro_heading = doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'This Employee Policy Manual ("Manual") provides guidelines and policies governing '
        'employment at Meridian Technologies, Inc. ("Company"). These policies are intended to '
        'create a productive, safe, and respectful workplace for all employees.'
    )
    doc.add_paragraph(
        'All employees are expected to read, understand, and comply with the policies outlined '
        'in this Manual. The Company reserves the right to modify, revise, or eliminate any '
        'policy at any time. Employees will be notified of significant changes.'
    )

    doc.add_page_break()

    # --- Page 2: Employment Policies ---
    doc.add_heading('Section 1: Employment Policies', level=1)
    doc.add_heading('1.1 Equal Employment Opportunity', level=2)
    doc.add_paragraph(
        'Meridian Technologies, Inc. is an equal opportunity employer. We are committed to '
        'providing employment opportunities without discrimination based on race, color, religion, '
        'sex, national origin, age, disability, veteran status, or any other characteristic '
        'protected by applicable law.'
    )
    doc.add_heading('1.2 At-Will Employment', level=2)
    doc.add_paragraph(
        'Employment at the Company is on an at-will basis. This means that either the employee '
        'or the Company may terminate the employment relationship at any time, with or without '
        'cause or advance notice, subject to applicable law.'
    )
    doc.add_heading('1.3 Background Checks', level=2)
    doc.add_paragraph(
        'All offers of employment are contingent upon the successful completion of a background '
        'check. The Company may conduct background checks as permitted by law, including criminal '
        'history, employment verification, and educational credential verification.'
    )

    doc.add_page_break()

    # --- Page 3: Code of Conduct ---
    doc.add_heading('Section 2: Code of Conduct', level=1)
    doc.add_heading('2.1 Professional Behavior', level=2)
    doc.add_paragraph(
        'Employees are expected to conduct themselves professionally at all times. This includes '
        'treating colleagues, customers, and vendors with respect and courtesy. Harassment, '
        'discrimination, or intimidation of any kind will not be tolerated.'
    )
    doc.add_heading('2.2 Conflict of Interest', level=2)
    doc.add_paragraph(
        'Employees must avoid situations that create actual or apparent conflicts of interest. '
        'A conflict of interest arises when an employee\'s personal interests interfere or '
        'appear to interfere with the interests of the Company. Employees must disclose any '
        'potential conflicts to their supervisor or Human Resources.'
    )
    doc.add_heading('2.3 Confidentiality', level=2)
    doc.add_paragraph(
        'All employees have a duty to protect the Company\'s confidential information. This '
        'includes but is not limited to proprietary technology, financial information, customer '
        'data, business strategies, and personnel records. Confidentiality obligations continue '
        'even after employment ends.'
    )

    doc.add_page_break()

    # --- Page 4: Work Hours and Attendance ---
    doc.add_heading('Section 3: Work Hours and Attendance', level=1)
    doc.add_heading('3.1 Standard Work Hours', level=2)
    doc.add_paragraph(
        'The standard work week consists of forty (40) hours, typically Monday through Friday, '
        '8:00 AM to 5:00 PM local time, with a one-hour lunch break. Individual departments '
        'may have adjusted schedules based on operational requirements.'
    )
    doc.add_heading('3.2 Attendance and Punctuality', level=2)
    doc.add_paragraph(
        'Regular, punctual attendance is essential to the effective operation of the Company. '
        'Employees are expected to report to work on time and to notify their supervisor as soon '
        'as possible if they will be absent or late. Unexcused absences or chronic tardiness '
        'may result in disciplinary action.'
    )
    doc.add_heading('3.3 Remote Work Policy', level=2)
    doc.add_paragraph(
        'Eligible employees may be permitted to work remotely on a full-time or part-time basis, '
        'subject to manager approval and departmental needs. Remote employees must maintain the '
        'same level of availability, productivity, and professionalism as in-office employees. '
        'Remote work arrangements may be revoked at any time.'
    )

    doc.add_page_break()

    # --- Page 5: Leave and Benefits ---
    doc.add_heading('Section 4: Leave and Benefits', level=1)
    doc.add_heading('4.1 Paid Time Off', level=2)
    doc.add_paragraph(
        'Full-time employees accrue fifteen (15) days of paid time off (PTO) per year during '
        'the first two years of employment. After two years, accrual increases to twenty (20) '
        'days per year. Part-time employees accrue PTO on a prorated basis. PTO must be '
        'approved in advance whenever possible.'
    )
    doc.add_heading('4.2 Holidays', level=2)
    doc.add_paragraph(
        'The Company observes ten (10) paid holidays per year, including New Year\'s Day, '
        'Memorial Day, Independence Day, Labor Day, Thanksgiving Day, the day after Thanksgiving, '
        'Christmas Eve, Christmas Day, and two floating holidays that employees may use at their '
        'discretion with manager approval.'
    )
    doc.add_heading('4.3 Health and Wellness Benefits', level=2)
    doc.add_paragraph(
        'The Company provides comprehensive health, dental, and vision insurance for all '
        'full-time employees and their eligible dependents. Employee contributions are deducted '
        'from bi-weekly paychecks on a pre-tax basis. Benefit enrollment information is '
        'provided during onboarding and during the annual open enrollment period.'
    )

    doc.add_page_break()

    # --- Page 6: Safety and Security ---
    doc.add_heading('Section 5: Workplace Safety and Security', level=1)
    doc.add_heading('5.1 Safety Commitment', level=2)
    doc.add_paragraph(
        'The health and safety of our employees is a top priority. The Company is committed to '
        'maintaining a safe and healthy work environment in compliance with all applicable '
        'occupational safety laws and regulations. Employees are expected to follow all safety '
        'guidelines and immediately report any unsafe conditions or accidents.'
    )
    doc.add_heading('5.2 Security Procedures', level=2)
    doc.add_paragraph(
        'Employees must wear their identification badge at all times while on Company premises. '
        'Access to secure areas is restricted to authorized personnel. Employees must not allow '
        'unauthorized individuals to enter secured areas. Lost or stolen badges must be reported '
        'to Security immediately.'
    )
    doc.add_heading('5.3 Drug-Free Workplace', level=2)
    doc.add_paragraph(
        'The Company maintains a drug-free workplace. The use, possession, distribution, or '
        'sale of illegal substances on Company premises or during Company activities is strictly '
        'prohibited. Employees may be subject to drug testing as permitted by law, including '
        'pre-employment, random, and post-incident testing.'
    )

    doc.add_page_break()

    # --- Page 7: Closing ---
    doc.add_heading('Section 6: Policy Acknowledgment', level=1)
    doc.add_paragraph(
        'By accepting employment with Meridian Technologies, Inc., employees acknowledge that '
        'they have received, read, and understood the policies contained in this Manual. '
        'Employees further acknowledge that this Manual does not constitute a contract of '
        'employment and that their employment remains at-will unless otherwise specified in a '
        'separate written agreement signed by an authorized Company representative.'
    )
    doc.add_paragraph(
        'This Policy Manual supersedes all previous policy manuals and any informal or verbal '
        'policies that may have been communicated. The Company reserves the right to add, '
        'change, or delete provisions of this Manual at any time without prior notice.'
    )
    doc.add_paragraph(
        'Employees are encouraged to ask their supervisor or Human Resources if they have any '
        'questions about the Company\'s policies or any matters not covered by this Manual. '
        'We are committed to maintaining clear and consistent communication with all employees.'
    )
    # This is the last paragraph that should be in the initial file
    doc.add_paragraph(
        'For questions regarding these policies, please contact the Human Resources department.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
