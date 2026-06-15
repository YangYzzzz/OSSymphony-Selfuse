"""
Initial Setup: Technical manual with glossary section containing 5 key terms in plain text
Task ID: osworld_writer_character_style_002
Domain: libreoffice_writer

Creates a technical networking manual with a glossary section.
The 5 key terms are in plain text (no 'KeyTerm' style, no bold, no special color).
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_character_style_002'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Network Infrastructure Technical Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph('Version 3.1 — Internal Reference Guide')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This technical manual provides a comprehensive reference for network '
        'infrastructure professionals and system administrators. It covers core '
        'concepts, operational guidelines, and troubleshooting procedures for '
        'enterprise-grade network environments.'
    )
    doc.add_paragraph(
        'The contents of this document are intended for personnel with at least '
        'an intermediate understanding of networking principles. Concepts are '
        'presented with practical examples drawn from real deployment scenarios.'
    )

    # --- Section 2: System Architecture ---
    doc.add_heading('2. System Architecture Overview', level=1)
    doc.add_paragraph(
        'Modern enterprise networks rely on a layered architecture that separates '
        'concerns across physical, data link, network, transport, and application '
        'layers. This separation allows teams to manage and troubleshoot each layer '
        'independently, reducing the complexity of large-scale deployments.'
    )

    doc.add_heading('2.1 Core Components', level=2)
    doc.add_paragraph(
        'The core of the network infrastructure consists of high-availability '
        'switches, routers, and load balancers. Each component is selected based '
        'on throughput requirements, redundancy needs, and integration capability '
        'with existing monitoring systems.'
    )
    doc.add_paragraph('Key hardware platforms in use include:', style='List Bullet')
    doc.add_paragraph('Cisco Catalyst 9000 series core switches', style='List Bullet')
    doc.add_paragraph('Juniper MX480 edge routers', style='List Bullet')
    doc.add_paragraph('F5 BIG-IP load balancers', style='List Bullet')
    doc.add_paragraph('Palo Alto PA-5000 firewalls', style='List Bullet')

    # --- Section 3: Performance Guidelines ---
    doc.add_heading('3. Performance Guidelines', level=1)
    doc.add_paragraph(
        'Network performance is measured across several dimensions including '
        'throughput, round-trip time, packet loss, and jitter. Establishing '
        'baseline measurements is essential before deploying new services or '
        'making configuration changes.'
    )
    doc.add_paragraph(
        'Scheduled maintenance windows should be used to perform firmware upgrades '
        'and configuration audits. All changes must be logged in the change '
        'management system and reviewed by a senior network engineer prior to '
        'implementation in production environments.'
    )

    doc.add_heading('3.1 Monitoring and Alerting', level=2)
    doc.add_paragraph(
        'Continuous monitoring using SNMP traps and NetFlow data allows the '
        'operations team to detect anomalies before they impact service quality. '
        'Alert thresholds should be configured conservatively to minimize false '
        'positives while ensuring genuine issues are escalated promptly.'
    )

    # --- Section 4: Troubleshooting ---
    doc.add_heading('4. Troubleshooting Procedures', level=1)
    doc.add_paragraph(
        'When responding to network incidents, engineers should follow a structured '
        'diagnostic approach. Begin by isolating the scope of the issue — whether '
        'it affects a single host, a subnet, or a broader network segment.'
    )
    doc.add_paragraph(
        'Common tools employed during diagnosis include ping, traceroute, nslookup, '
        'Wireshark packet capture, and platform-specific CLI commands. Document all '
        'findings in the incident management system throughout the investigation.'
    )

    # --- Section 5: Glossary ---
    doc.add_heading('5. Glossary', level=1)
    doc.add_paragraph(
        'The following terms appear throughout this document. Definitions are '
        'provided here for quick reference.'
    )

    # 5 key terms — plain text, no special style, no bold, no color, default font size
    glossary_items = [
        (
            'Application Programming Interface (API)',
            'A set of rules and protocols that allows different software applications '
            'to communicate with each other. APIs define the methods and data formats '
            'that applications can use to request and exchange information.'
        ),
        (
            'Bandwidth',
            'The maximum rate of data transfer across a network path, typically '
            'measured in megabits per second (Mbps) or gigabits per second (Gbps). '
            'Higher bandwidth allows more data to be transmitted simultaneously.'
        ),
        (
            'Cache',
            'A temporary storage location that holds copies of frequently accessed '
            'data to reduce retrieval time. Caching is used at multiple layers, '
            'including CPU caches, browser caches, and content delivery network caches.'
        ),
        (
            'DNS (Domain Name System)',
            'A hierarchical and decentralized naming system that translates human-readable '
            'domain names (such as www.example.com) into numerical IP addresses that '
            'network devices use to locate and connect to each other.'
        ),
        (
            'Latency',
            'The time delay between a request being made and the corresponding response '
            'being received, typically measured in milliseconds. Low latency is critical '
            'for real-time applications such as voice communications and video conferencing.'
        ),
    ]

    for term, definition in glossary_items:
        # Term on its own line — plain text, no bold, no color, no special style
        term_para = doc.add_paragraph()
        term_run = term_para.add_run(term)
        # Explicitly ensure no special formatting (plain)
        term_run.bold = False
        term_run.font.color.rgb = None  # default color

        # Definition indented as next paragraph
        def_para = doc.add_paragraph(definition)
        def_para.paragraph_format.left_indent = Pt(20)

        doc.add_paragraph('')  # spacer between glossary entries

    # --- Appendix ---
    doc.add_heading('Appendix A: Reference Standards', level=1)
    doc.add_paragraph(
        'This document references the following industry standards and RFCs: '
        'RFC 1035 (DNS), RFC 2616 (HTTP/1.1), RFC 7540 (HTTP/2), '
        'IEEE 802.1Q (VLAN tagging), IEEE 802.3 (Ethernet).'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
