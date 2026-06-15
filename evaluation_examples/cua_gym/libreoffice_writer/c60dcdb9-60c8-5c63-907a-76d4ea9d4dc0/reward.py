"""
Reward Script: Create 'KeyTerm' character style and apply to 5 glossary terms
Task ID: osworld_writer_character_style_002
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4 pts): 'KeyTerm' character style exists with correct properties
                           (bold, dark red #8B0000, 11pt)
  Component 2 (0.6 pts): All 5 glossary key term paragraphs have the 'KeyTerm'
                           style applied (bold=True, color=8B0000, size=11pt)
  Total: 1.0
"""

import os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_character_style_002'

# Expected formatting for the KeyTerm character style
EXPECTED_COLOR = RGBColor(0x8B, 0x00, 0x00)  # dark red #8B0000
EXPECTED_SIZE_PT = 11.0
EXPECTED_BOLD = True  # the style requires bold

# The 5 glossary key term paragraph indices (0-based) in the document
# Para 25: Application Programming Interface (API)
# Para 28: Bandwidth
# Para 31: Cache
# Para 34: DNS (Domain Name System)
# Para 37: Latency
GLOSSARY_TERM_INDICES = [25, 28, 31, 34, 37]

# Color tolerance for comparison (max channel difference)
COLOR_TOLERANCE = 10


def colors_close(actual_rgb, expected_rgb, tolerance=COLOR_TOLERANCE):
    """Check if two RGBColor values are within tolerance."""
    try:
        diff = (
            abs(actual_rgb[0] - expected_rgb[0]) +
            abs(actual_rgb[1] - expected_rgb[1]) +
            abs(actual_rgb[2] - expected_rgb[2])
        )
        return diff <= tolerance * 3
    except Exception:
        return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: 'KeyTerm' character style exists with correct properties (0.4 pts)
    # The task requires creating a NEW character style named 'KeyTerm' with:
    #   - bold formatting
    #   - dark red color (#8B0000)
    #   - 11pt font size
    # -----------------------------------------------------------------------
    try:
        keyterm_style = None
        for style in doc.styles:
            if style.name == 'KeyTerm':
                keyterm_style = style
                break

        if keyterm_style is None:
            print("FAIL: Component 1 — 'KeyTerm' character style not found in document")
        elif keyterm_style.type != WD_STYLE_TYPE.CHARACTER:
            print(f"FAIL: Component 1 — 'KeyTerm' style found but type is {keyterm_style.type}, expected CHARACTER")
        else:
            # Verify style font properties
            style_font = keyterm_style.font
            style_bold = style_font.bold
            style_size_pt = style_font.size.pt if style_font.size else None
            style_color = style_font.color.rgb if (style_font.color and style_font.color.type) else None

            checks = []

            # Bold check
            if style_bold is True:
                checks.append(("bold", True))
                print(f"PASS: Component 1a — KeyTerm style bold=True")
            else:
                checks.append(("bold", False))
                print(f"FAIL: Component 1a — KeyTerm style bold={style_bold}, expected True")

            # Size check
            if style_size_pt is not None and abs(style_size_pt - EXPECTED_SIZE_PT) < 0.5:
                checks.append(("size", True))
                print(f"PASS: Component 1b — KeyTerm style size={style_size_pt}pt")
            else:
                checks.append(("size", False))
                print(f"FAIL: Component 1b — KeyTerm style size={style_size_pt}pt, expected {EXPECTED_SIZE_PT}pt")

            # Color check
            if style_color is not None and colors_close(style_color, EXPECTED_COLOR):
                checks.append(("color", True))
                print(f"PASS: Component 1c — KeyTerm style color={style_color} (expected #8B0000)")
            else:
                checks.append(("color", False))
                print(f"FAIL: Component 1c — KeyTerm style color={style_color}, expected {EXPECTED_COLOR}")

            all_pass = all(v for _, v in checks)
            if all_pass:
                print(f"PASS: Component 1 — 'KeyTerm' character style with correct bold, 11pt, #8B0000 (0.4 pts)")
                total_score += 0.4
            else:
                # Give partial credit if style exists but some properties are wrong
                pass_count = sum(1 for _, v in checks if v)
                if pass_count > 0:
                    partial = round(0.4 * pass_count / 3, 4)
                    print(f"PARTIAL: Component 1 — KeyTerm style exists but {3-pass_count}/3 property checks failed ({partial} pts)")
                    total_score += partial
                else:
                    print(f"FAIL: Component 1 — 'KeyTerm' style exists but has wrong properties (0.0 pts)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: All 5 glossary key terms have the 'KeyTerm' style applied (0.6 pts)
    # Each of the 5 glossary term paragraphs must have runs with:
    #   - run.style.name == 'KeyTerm'  (character style applied)
    # AND the effective formatting should show:
    #   - bold=True, color=8B0000, size=11pt
    # We award 0.12 pts per term (5 terms × 0.12 = 0.60)
    # -----------------------------------------------------------------------
    try:
        num_paras = len(doc.paragraphs)
        terms_styled_correctly = 0
        term_names = [
            'Application Programming Interface (API)',
            'Bandwidth',
            'Cache',
            'DNS (Domain Name System)',
            'Latency',
        ]

        for i, para_idx in enumerate(GLOSSARY_TERM_INDICES):
            term_name = term_names[i]
            if para_idx >= num_paras:
                print(f"FAIL: Component 2 — Para index {para_idx} out of range (doc has {num_paras} paras)")
                continue

            para = doc.paragraphs[para_idx]

            # Check that all non-empty runs in this paragraph have the KeyTerm style
            # and the expected formatting
            runs = [r for r in para.runs if r.text.strip()]
            if not runs:
                print(f"FAIL: Component 2 — Para {para_idx} ({term_name!r}) has no text runs")
                continue

            term_ok = True
            for run in runs:
                run_style_name = run.style.name if run.style else None
                run_bold = run.font.bold
                run_size_pt = run.font.size.pt if run.font.size else None
                run_color = run.font.color.rgb if (run.font.color and run.font.color.type) else None

                style_applied = (run_style_name == 'KeyTerm')
                bold_ok = (run_bold is True)
                size_ok = (run_size_pt is not None and abs(run_size_pt - EXPECTED_SIZE_PT) < 0.5)
                color_ok = (run_color is not None and colors_close(run_color, EXPECTED_COLOR))

                if not (style_applied and bold_ok and size_ok and color_ok):
                    term_ok = False
                    print(
                        f"FAIL: Component 2 — Para {para_idx} ({term_name!r}) run formatting: "
                        f"style={run_style_name!r}(ok={style_applied}), "
                        f"bold={run_bold}(ok={bold_ok}), "
                        f"size_pt={run_size_pt}(ok={size_ok}), "
                        f"color={run_color}(ok={color_ok})"
                    )
                    break

            if term_ok:
                print(f"PASS: Component 2 — Para {para_idx} ({term_name!r}) correctly styled with KeyTerm (0.12 pts)")
                terms_styled_correctly += 1

        comp2_score = round(terms_styled_correctly * 0.12, 4)
        total_score += comp2_score
        if terms_styled_correctly == 5:
            print(f"PASS: Component 2 — All 5 glossary terms styled with KeyTerm ({comp2_score} pts)")
        else:
            print(f"PARTIAL: Component 2 — {terms_styled_correctly}/5 glossary terms styled correctly ({comp2_score} pts)")

    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM environment
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
