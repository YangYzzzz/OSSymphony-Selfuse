"""
Initial Setup: Plain-text attendance record for 15 employees
Task ID: writer_hr_046
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_046'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# Employee attendance data: Name, Department, Days Present, Days Absent, Attendance Rate
EMPLOYEES = [
    ("Sarah Chen", "Engineering", 228, 12, 95.0),
    ("Marcus Johnson", "Marketing", 210, 30, 87.5),
    ("Priya Patel", "Finance", 232, 8, 96.7),
    ("David Kim", "Engineering", 205, 35, 85.4),
    ("Elena Rodriguez", "Human Resources", 235, 5, 97.9),
    ("James O'Brien", "Sales", 220, 20, 91.7),
    ("Aisha Mohammed", "Finance", 238, 2, 99.2),
    ("Robert Taylor", "Marketing", 198, 42, 82.5),
    ("Lin Wei", "Engineering", 230, 10, 95.8),
    ("Sophie Martin", "Sales", 212, 28, 88.3),
    ("Carlos Gutierrez", "Operations", 225, 15, 93.8),
    ("Fatima Al-Rashid", "Human Resources", 215, 25, 89.6),
    ("Thomas Anderson", "Operations", 233, 7, 97.1),
    ("Yuki Tanaka", "Finance", 208, 32, 86.7),
    ("Michael Foster", "Sales", 222, 18, 92.5),
]


def create_initial():
    doc = Document()

    # Add a simple title paragraph (unformatted, just plain text)
    title_para = doc.add_paragraph("Attendance Record")
    title_run = title_para.runs[0]
    title_run.font.size = Pt(16)

    # Add blank line
    doc.add_paragraph("")

    # Add header line as plain text with tabs
    header_line = "Employee Name\tDepartment\tDays Present\tDays Absent\tAttendance Rate"
    doc.add_paragraph(header_line)

    # Add each employee as a plain-text tab-separated line
    for name, dept, present, absent, rate in EMPLOYEES:
        line = f"{name}\t{dept}\t{present}\t{absent}\t{rate}%"
        doc.add_paragraph(line)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
