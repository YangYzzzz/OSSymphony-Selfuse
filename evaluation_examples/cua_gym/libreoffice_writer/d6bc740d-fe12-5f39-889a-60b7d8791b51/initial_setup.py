"""
Initial Setup: Create photo_fix_list.docx and old_photo.jpg for GIMP restoration task.
Task ID: osworld_multi_apps_writer_gimp_067
Domain: libreoffice_writer + gimp (multi-app)

Creates:
  - /home/user/photo_fix_list.docx  (correction list open in LibreOffice Writer)
  - /home/user/Desktop/old_photo.jpg (vintage scanned photo with scratch at y=320)
"""

import os
import shlex
import subprocess
import time
import numpy as np

WORKDIR = '/home/user'
DESKTOP = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_writer_gimp_067'
DOC_OUTPUT = f'{WORKDIR}/photo_fix_list.docx'
PHOTO_OUTPUT = f'{DESKTOP}/old_photo.jpg'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_fix_list_doc():
    """Create the photo_fix_list.docx with correction instructions."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Title
    title = doc.add_heading("Photo Restoration Fix List", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle / intro
    intro = doc.add_paragraph(
        "The following corrections must be applied to 'old_photo.jpg' using GIMP. "
        "Please complete each step in order and save the final result as 'old_photo_restored.jpg' on the Desktop."
    )
    intro.paragraph_format.space_after = Pt(12)

    doc.add_paragraph("")  # spacer

    # Step 1
    step1_heading = doc.add_heading("Step 1: Remove Scratch", level=2)
    step1 = doc.add_paragraph()
    run = step1.add_run("Use the ")
    run2 = step1.add_run("Heal tool")
    run2.bold = True
    run3 = step1.add_run(
        " to remove the diagonal scratch at approximately y=320 (horizontal band). "
        "The scratch runs across the full width of the image near row 320."
    )

    doc.add_paragraph("")  # spacer

    # Step 2
    step2_heading = doc.add_heading("Step 2: Exposure Correction", level=2)
    step2 = doc.add_paragraph()
    run4 = step2.add_run("Use ")
    run5 = step2.add_run("Colors > Curves")
    run5.bold = True
    run6 = step2.add_run(
        " to increase the exposure. Raise the midpoints of the curve by "
    )
    run7 = step2.add_run("15 units")
    run7.bold = True
    run8 = step2.add_run(
        " to brighten the midtones and recover the underexposed areas of the vintage photograph."
    )

    doc.add_paragraph("")  # spacer

    # Step 3
    step3_heading = doc.add_heading("Step 3: Convert to Greyscale", level=2)
    step3 = doc.add_paragraph()
    run9 = step3.add_run("Convert the image to ")
    run10 = step3.add_run("greyscale (Image > Mode > Greyscale)")
    run10.bold = True
    run11 = step3.add_run(
        " to give it an authentic vintage monochromatic appearance consistent with old photographs."
    )

    doc.add_paragraph("")  # spacer

    # Output instructions
    output_heading = doc.add_heading("Output", level=2)
    output_para = doc.add_paragraph()
    run12 = output_para.add_run("Save the restored image as ")
    run13 = output_para.add_run("'old_photo_restored.jpg'")
    run13.bold = True
    run14 = output_para.add_run(
        " on the Desktop. Use File > Export As in GIMP to save as JPEG format."
    )

    doc.add_paragraph("")  # spacer

    # Footer note
    note = doc.add_paragraph(
        "Note: The original 'old_photo.jpg' must remain unmodified. "
        "All corrections should be applied to a working copy and saved under the new filename."
    )
    note.paragraph_format.space_before = Pt(6)
    for run in note.runs:
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.size = Pt(9)

    doc.save(DOC_OUTPUT)
    print(f'Fix list document created: {DOC_OUTPUT}')


def create_vintage_photo():
    """Create a realistic vintage-looking scanned photo with a scratch at y=320."""
    from PIL import Image, ImageDraw, ImageFilter, ImageEnhance
    import numpy as np

    width, height = 800, 600
    rng = np.random.default_rng(42)

    # --- Base: create a sepia-toned vintage landscape ---
    # Start with a sky gradient (top third)
    img_array = np.zeros((height, width, 3), dtype=np.uint8)

    # Sky region (top 200px): light grey-blue sky
    for y in range(200):
        brightness = int(180 + (y / 200) * 30)
        img_array[y, :] = [brightness, brightness - 10, brightness - 20]

    # Horizon / hills region (200–350): dark green hills
    for y in range(200, 350):
        t = (y - 200) / 150
        r = int(80 + t * 60)
        g = int(90 + t * 50)
        b = int(60 + t * 30)
        # add some texture variation across columns
        noise = rng.integers(-15, 15, width)
        img_array[y, :, 0] = np.clip(r + noise, 0, 255)
        img_array[y, :, 1] = np.clip(g + noise, 0, 255)
        img_array[y, :, 2] = np.clip(b + noise, 0, 255)

    # Ground region (350–600): brownish ground with texture
    for y in range(350, height):
        t = (y - 350) / (height - 350)
        r = int(120 + t * 40)
        g = int(100 + t * 30)
        b = int(70 + t * 20)
        noise = rng.integers(-20, 20, width)
        img_array[y, :, 0] = np.clip(r + noise, 0, 255)
        img_array[y, :, 1] = np.clip(g + noise, 0, 255)
        img_array[y, :, 2] = np.clip(b + noise, 0, 255)

    img = Image.fromarray(img_array, 'RGB')

    # Add a simple tree silhouette
    draw = ImageDraw.Draw(img)
    # Tree trunk
    draw.rectangle([380, 290, 400, 370], fill=(60, 45, 30))
    # Tree canopy (dark green)
    draw.ellipse([330, 210, 450, 310], fill=(40, 70, 40))
    draw.ellipse([345, 195, 440, 290], fill=(50, 80, 45))

    # Add a small farmhouse silhouette
    # House body
    draw.rectangle([550, 300, 650, 370], fill=(100, 80, 60))
    # Roof
    draw.polygon([(540, 300), (660, 300), (600, 255)], fill=(70, 50, 40))
    # Window
    draw.rectangle([570, 320, 595, 345], fill=(180, 160, 100))
    draw.rectangle([620, 320, 645, 345], fill=(180, 160, 100))
    # Door
    draw.rectangle([593, 345, 615, 370], fill=(60, 40, 25))

    # Foreground fence posts
    for x in range(50, 750, 60):
        draw.rectangle([x, 360, x + 6, 420], fill=(90, 70, 50))
    # Fence wire
    draw.line([(50, 380), (750, 380)], fill=(80, 65, 45), width=2)
    draw.line([(50, 400), (750, 400)], fill=(80, 65, 45), width=2)

    # Apply sepia tone
    img_array2 = np.array(img).astype(np.float32)
    r = img_array2[:, :, 0]
    g = img_array2[:, :, 1]
    b = img_array2[:, :, 2]
    sepia_r = np.clip(r * 0.393 + g * 0.769 + b * 0.189, 0, 255)
    sepia_g = np.clip(r * 0.349 + g * 0.686 + b * 0.168, 0, 255)
    sepia_b = np.clip(r * 0.272 + g * 0.534 + b * 0.131, 0, 255)
    sepia_array = np.stack([sepia_r, sepia_g, sepia_b], axis=2).astype(np.uint8)
    img = Image.fromarray(sepia_array, 'RGB')

    # Make slightly darker (underexposed old photo look)
    enhancer = ImageEnhance.Brightness(img)
    img = enhancer.enhance(0.82)

    # Add film grain (noise texture)
    img_array3 = np.array(img).astype(np.int16)
    grain = rng.integers(-18, 18, (height, width, 3), dtype=np.int16)
    img_array3 = np.clip(img_array3 + grain, 0, 255).astype(np.uint8)
    img = Image.fromarray(img_array3, 'RGB')

    # Apply slight blur for aged look
    img = img.filter(ImageFilter.GaussianBlur(radius=0.5))

    # Add slight vignette
    img_array4 = np.array(img).astype(np.float32)
    Y, X = np.ogrid[:height, :width]
    cx, cy = width / 2, height / 2
    dist = np.sqrt(((X - cx) / cx) ** 2 + ((Y - cy) / cy) ** 2)
    vignette = 1.0 - 0.35 * dist
    vignette = np.clip(vignette, 0, 1)
    for c in range(3):
        img_array4[:, :, c] *= vignette
    img = Image.fromarray(np.clip(img_array4, 0, 255).astype(np.uint8), 'RGB')

    # --- ADD THE SCRATCH at y=320 ---
    # The scratch is a bright/light diagonal streak near row 320
    img_array5 = np.array(img)
    # Scratch: a thin near-horizontal line with slight wiggle, mostly white/bright
    scratch_y = 320
    scratch_color = 220  # bright highlight
    for x in range(width):
        # Slight vertical wobble
        wobble = int(2 * np.sin(x * 0.05))
        y_pos = scratch_y + wobble
        # Main scratch line (1-2px wide)
        for dy in range(2):
            y_idx = y_pos + dy
            if 0 <= y_idx < height:
                img_array5[y_idx, x] = [scratch_color, scratch_color - 5, scratch_color - 10]
        # Secondary faint scratch line (partial)
        if x % 3 != 0:
            y_idx2 = y_pos + 3
            if 0 <= y_idx2 < height:
                curr = img_array5[y_idx2, x].astype(np.int16)
                brightened = np.clip(curr + 40, 0, 255).astype(np.uint8)
                img_array5[y_idx2, x] = brightened

    img = Image.fromarray(img_array5, 'RGB')

    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    # Save as JPEG (vintage scanned photo quality)
    img.save(PHOTO_OUTPUT, 'JPEG', quality=88)
    print(f'Vintage photo created: {PHOTO_OUTPUT}')
    print(f'  Size: {img.size}, Mode: {img.mode}')
    print(f'  Scratch added at y~{scratch_y}')

    # Ensure the restored file does NOT pre-exist
    restored_path = f'{DESKTOP}/old_photo_restored.jpg'
    if os.path.exists(restored_path):
        os.remove(restored_path)
        print(f'Removed pre-existing restored file: {restored_path}')


def main():
    create_fix_list_doc()
    create_vintage_photo()

    # GUI-ready startup:
    # 1. Open photo_fix_list.docx in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DOC_OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with photo_fix_list.docx')


main()
