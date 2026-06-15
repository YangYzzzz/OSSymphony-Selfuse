"""
Initial Setup: CloudSync API Documentation - title formatting task
Task ID: writer_tech_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title (regular weight, 12pt, left-aligned) ---
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_run = title_para.add_run("CloudSync API Documentation v2.1")
    title_run.bold = False
    title_run.font.size = Pt(12)
    title_run.font.name = "Calibri"

    # --- Section: Overview ---
    overview_heading = doc.add_heading("Overview", level=2)

    p1 = doc.add_paragraph(
        "CloudSync is a RESTful API service designed for seamless file "
        "synchronization across multiple cloud storage providers. This "
        "documentation covers authentication, endpoint specifications, "
        "rate limiting, and error handling for version 2.1 of the API."
    )

    # --- Section: Authentication ---
    auth_heading = doc.add_heading("Authentication", level=2)

    p2 = doc.add_paragraph(
        "All API requests must include a valid Bearer token in the "
        "Authorization header. Tokens are obtained through the OAuth 2.0 "
        "client credentials flow via the /auth/token endpoint."
    )

    p3 = doc.add_paragraph(
        "Token expiration is set to 3600 seconds by default. Refresh tokens "
        "remain valid for 30 days and can be rotated using the /auth/refresh "
        "endpoint."
    )

    # --- Section: Endpoints ---
    endpoints_heading = doc.add_heading("Core Endpoints", level=2)

    # Endpoint table
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ["Method", "Endpoint", "Description"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    endpoints = [
        ["GET", "/api/v2/files", "List all synced files"],
        ["POST", "/api/v2/files/upload", "Upload a new file"],
        ["PUT", "/api/v2/files/{id}/sync", "Trigger sync for a specific file"],
        ["DELETE", "/api/v2/files/{id}", "Remove a file from sync"],
        ["PATCH", "/api/v2/files/{id}/metadata", "Update file metadata"],
    ]
    for r, row_data in enumerate(endpoints, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Section: Rate Limits ---
    rate_heading = doc.add_heading("Rate Limits", level=2)

    p4 = doc.add_paragraph(
        "Standard tier accounts are limited to 1,000 requests per minute. "
        "Enterprise accounts receive a quota of 10,000 requests per minute. "
        "When the rate limit is exceeded, the API returns HTTP 429 with a "
        "Retry-After header indicating the cooldown period in seconds."
    )

    # --- Section: Error Codes ---
    error_heading = doc.add_heading("Error Codes", level=2)

    error_table = doc.add_table(rows=5, cols=3)
    error_table.style = "Table Grid"
    error_headers = ["Code", "Status", "Description"]
    for i, h in enumerate(error_headers):
        cell = error_table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    errors = [
        ["400", "Bad Request", "Malformed request body or missing parameters"],
        ["401", "Unauthorized", "Invalid or expired authentication token"],
        ["404", "Not Found", "Requested resource does not exist"],
        ["429", "Too Many Requests", "Rate limit exceeded"],
    ]
    for r, row_data in enumerate(errors, 1):
        for c, val in enumerate(row_data):
            error_table.cell(r, c).text = val

    # --- Section: Changelog ---
    changelog_heading = doc.add_heading("Changelog", level=2)

    p5 = doc.add_paragraph(
        "v2.1 (2025-09-15): Added PATCH endpoint for metadata updates. "
        "Improved error response formatting. Increased enterprise rate limit "
        "from 5,000 to 10,000 requests per minute."
    )
    p6 = doc.add_paragraph(
        "v2.0 (2025-03-01): Major refactor of authentication flow. Migrated "
        "to OAuth 2.0. Deprecated legacy API key authentication."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
