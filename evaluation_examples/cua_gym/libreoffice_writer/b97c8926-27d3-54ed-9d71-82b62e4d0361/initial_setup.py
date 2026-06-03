"""
Initial Setup: sidebar_layout.docx - 3-page document with 2 equal columns, no separator
Task ID: writer_page_052
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'sidebar_layout'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_two_equal_columns(section):
    """Set section to 2 equal columns with 0.5cm spacing, no separator."""
    sectPr = section._sectPr
    # Remove any existing cols element
    for col_elem in sectPr.findall(qn('w:cols')):
        sectPr.remove(col_elem)

    # Create w:cols with 2 equal columns
    # w:num=2, w:space in twips (0.5cm = 283 twips), equalWidth=1 (no separator)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '283')      # 0.5cm in twips (rounded)
    cols.set(qn('w:equalWidth'), '1')
    # NO w:sep attribute (no separator line)

    # Insert before docGrid if present, otherwise append
    doc_grid = sectPr.find(qn('w:docGrid'))
    if doc_grid is not None:
        sectPr.insert(list(sectPr).index(doc_grid), cols)
    else:
        sectPr.append(cols)


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Page setup: A4 portrait ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # --- Column layout: 2 equal columns, 0.5cm spacing, NO separator ---
    set_two_equal_columns(section)

    # --- Page 1: Introduction ---
    h1 = doc.add_heading('Project Overview', level=1)

    p1 = doc.add_paragraph(
        'This document provides a comprehensive overview of the annual product strategy '
        'initiative. The marketing team has identified three key growth areas for the '
        'coming fiscal year: digital expansion, partnership development, and customer '
        'retention programs.'
    )

    p2 = doc.add_paragraph(
        'Market analysis conducted in Q3 revealed significant opportunities in the '
        'enterprise segment. Revenue projections indicate a potential 23% increase if '
        'the outlined strategies are implemented by Q2 of next year.'
    )

    sidebar1 = doc.add_paragraph(
        'Key Metric: Current customer acquisition cost stands at $142 per user. '
        'Target reduction to $95 through improved onboarding funnel optimization.'
    )
    sidebar1.paragraph_format.left_indent = Cm(0)

    p3 = doc.add_paragraph(
        'The competitive landscape has shifted considerably following the recent merger '
        'of two major industry players. Our positioning strategy emphasizes unique value '
        'propositions that differentiate our offering from commoditized alternatives.'
    )

    p4 = doc.add_paragraph(
        'Product development roadmap includes five major feature releases scheduled '
        'between January and June. Each release targets specific user segments identified '
        'in the customer research conducted by the UX team during September.'
    )

    # Page break to page 2
    doc.add_page_break()

    # --- Page 2: Financial Analysis ---
    h2 = doc.add_heading('Financial Analysis', level=1)

    p5 = doc.add_paragraph(
        'Budget allocation for the upcoming year has been finalized after extensive '
        'review by the finance committee. Total approved budget: $4.2 million, '
        'representing an 18% increase from the previous fiscal year allocation.'
    )

    p6 = doc.add_paragraph(
        'Engineering department receives the largest share at 35% ($1.47M), followed '
        'by sales and marketing at 28% ($1.18M). Operations and infrastructure '
        'accounts for 22% ($924K), with the remaining 15% allocated to administration.'
    )

    sidebar2 = doc.add_paragraph(
        'Note: Q1 spending projections assume full headcount by February 1st. '
        'Hiring delays may result in underspend of up to $180K in personnel costs.'
    )

    p7 = doc.add_paragraph(
        'Revenue recognition policy changes effective January 1st will require '
        'adjustments to how subscription contracts are reported. The accounting team '
        'has prepared transition documentation outlining the new procedures.'
    )

    p8 = doc.add_paragraph(
        'Investment in tooling and infrastructure upgrades totals $320K. Primary '
        'expenditures include cloud infrastructure scaling ($145K), security compliance '
        'tools ($95K), and developer productivity platforms ($80K).'
    )

    p9 = doc.add_paragraph(
        'Cost reduction initiatives identified by the operations team are expected to '
        'yield $230K in annual savings. Process automation projects account for the '
        'majority of savings through reduced manual processing time.'
    )

    # Page break to page 3
    doc.add_page_break()

    # --- Page 3: Implementation Timeline ---
    h3 = doc.add_heading('Implementation Timeline', level=1)

    p10 = doc.add_paragraph(
        'Phase 1 (January–March): Foundation work including infrastructure setup, '
        'team onboarding, and initial product development sprints. All team leads '
        'must submit detailed execution plans by December 15th.'
    )

    p11 = doc.add_paragraph(
        'Phase 2 (April–June): Feature rollout and beta testing program. Target '
        'audience: 500 enterprise customers selected from the existing customer base. '
        'Feedback collection through structured interviews and usage analytics.'
    )

    sidebar3 = doc.add_paragraph(
        'Milestone: Beta program launch requires legal review completion and '
        'data processing agreement updates. Estimated legal review duration: 3 weeks.'
    )

    p12 = doc.add_paragraph(
        'Phase 3 (July–September): General availability launch and scale-up. '
        'Marketing campaigns activate across all channels simultaneously. Customer '
        'success team expands by 8 additional headcount to support growth.'
    )

    p13 = doc.add_paragraph(
        'Phase 4 (October–December): Optimization and consolidation. Performance '
        'metrics review, strategy adjustment based on market feedback, and planning '
        'initiation for the following fiscal year objectives.'
    )

    p14 = doc.add_paragraph(
        'Risk mitigation plan covers supply chain dependencies, regulatory changes, '
        'and competitive responses. Contingency budget of $180K reserved for '
        'unplanned expenditures arising from market shifts or technical debt resolution.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
