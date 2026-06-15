"""
Reward Script: Apply 'Legal Body Text' custom paragraph style to all body paragraphs
Task ID: writer_legal_030
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15): 'Legal Body Text' style exists in document
  Component 2 (0.15): Style font is Times New Roman 12pt
  Component 3 (0.15): Style alignment is JUSTIFY
  Component 4 (0.15): Style line spacing is 1.5
  Component 5 (0.15): Style space_after is 6pt and first_line_indent is 0.5 inches
  Component 6 (0.15): All body paragraphs use 'Legal Body Text' style
  Component 7 (0.10): Headings and Title retain their original styles
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_030'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: 'Legal Body Text' style exists (0.15 points)
    try:
        style_names = [s.name for s in doc.styles]
        if 'Legal Body Text' in style_names:
            print(f"PASS: Component 1 - 'Legal Body Text' style exists (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 - 'Legal Body Text' style not found. Available paragraph styles: {[s.name for s in doc.styles if s.type and s.type.name == 'PARAGRAPH']}")
            # If the style doesn't exist, no further style checks can pass
            final_score = min(total_score, 1.0)
            print(f"\nScore: {total_score}/1.0")
            print(f"REWARD: {final_score}")
            return final_score
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Get the style object for remaining checks
    try:
        lbt_style = doc.styles['Legal Body Text']
    except Exception as e:
        print(f"ERROR: Cannot access 'Legal Body Text' style: {e}")
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Component 2: Style font is Times New Roman 12pt (0.15 points)
    try:
        font = lbt_style.font
        font_name = font.name
        font_size_pt = font.size.pt if font.size else None
        if font_name == 'Times New Roman' and font_size_pt == 12.0:
            print(f"PASS: Component 2 - Font is Times New Roman 12pt (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 - Expected Times New Roman 12pt, found name='{font_name}', size={font_size_pt}pt")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Style alignment is JUSTIFY (0.15 points)
    try:
        pf = lbt_style.paragraph_format
        alignment = pf.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            print(f"PASS: Component 3 - Alignment is JUSTIFY (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 - Expected JUSTIFY alignment, found {alignment}")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Style line spacing is 1.5 (0.15 points)
    try:
        pf = lbt_style.paragraph_format
        line_spacing = pf.line_spacing
        if line_spacing is not None and abs(float(line_spacing) - 1.5) < 0.01:
            print(f"PASS: Component 4 - Line spacing is 1.5 (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 4 - Expected line spacing 1.5, found {line_spacing}")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Style space_after is 6pt and first_line_indent is 0.5 inches (0.15 points)
    try:
        pf = lbt_style.paragraph_format
        space_after = pf.space_after
        first_line_indent = pf.first_line_indent

        space_after_pt = space_after.pt if space_after else None
        # Convert first_line_indent from EMU to inches (1 inch = 914400 EMU)
        fli_inches = first_line_indent / 914400 if first_line_indent else None

        space_ok = space_after_pt is not None and abs(space_after_pt - 6.0) < 0.1
        fli_ok = fli_inches is not None and abs(fli_inches - 0.5) < 0.01

        if space_ok and fli_ok:
            print(f"PASS: Component 5 - Space after=6pt, first line indent=0.5in (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 - Expected space_after=6pt (got {space_after_pt}), first_line_indent=0.5in (got {fli_inches})")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: All body paragraphs use 'Legal Body Text' style (0.15 points)
    # Body paragraphs = those that are NOT Title, Heading, or empty
    try:
        body_paras = []
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else 'Normal'
            # Skip Title, Headings
            if style_name == 'Title' or style_name.startswith('Heading'):
                continue
            # This is a body paragraph — should use Legal Body Text
            body_paras.append((para.text[:60], style_name))

        if len(body_paras) == 0:
            print(f"FAIL: Component 6 - No body paragraphs found")
        else:
            lbt_count = sum(1 for _, sn in body_paras if sn == 'Legal Body Text')
            total_body = len(body_paras)
            if lbt_count == total_body:
                print(f"PASS: Component 6 - All {total_body} body paragraphs use 'Legal Body Text' style (0.15 pts)")
                total_score += 0.15
            else:
                wrong = [(t, sn) for t, sn in body_paras if sn != 'Legal Body Text']
                print(f"FAIL: Component 6 - {lbt_count}/{total_body} body paragraphs use 'Legal Body Text'. Wrong: {wrong[:3]}")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    # Component 7: Headings and Title retain their original styles (0.10 points)
    try:
        heading_paras = [(i, para.style.name, para.text[:60]) for i, para in enumerate(doc.paragraphs)
                         if para.style.name == 'Title' or para.style.name.startswith('Heading')]
        # We expect: P0=Title, and multiple Heading 1 paragraphs
        has_title = any(sn == 'Title' for _, sn, _ in heading_paras)
        has_headings = any(sn.startswith('Heading') for _, sn, _ in heading_paras)
        if has_title and has_headings:
            print(f"PASS: Component 7 - Title and Headings retain their styles ({len(heading_paras)} non-body paragraphs) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 7 - Expected Title and Heading styles. Found: {heading_paras}")
    except Exception as e:
        print(f"ERROR: Component 7 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
