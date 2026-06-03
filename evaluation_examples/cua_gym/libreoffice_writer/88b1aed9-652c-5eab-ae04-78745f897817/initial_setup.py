"""
Initial Setup: Legal agreement with default paragraph styling
Task ID: writer_legal_030
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_030'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default style to Liberation Serif 12pt, single spacing, no indent
    style = doc.styles['Normal']
    style.font.name = 'Liberation Serif'
    style.font.size = Pt(12)
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.first_line_indent = Inches(0)

    # --- Document Title ---
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)

    # --- Preamble ---
    doc.add_paragraph(
        'This Professional Services Agreement ("Agreement") is entered into as of '
        'March 15, 2025, by and between Meridian Consulting Group, LLC, a Delaware '
        'limited liability company with its principal place of business at 4200 Corporate '
        'Plaza Drive, Suite 800, Chicago, Illinois 60601 ("Consultant"), and Westfield '
        'Technologies, Inc., a California corporation with its principal place of business '
        'at 1750 Innovation Boulevard, San Jose, California 95134 ("Client").'
    )

    doc.add_paragraph(
        'WHEREAS, the Client desires to engage the Consultant to provide certain professional '
        'consulting services as described herein; and WHEREAS, the Consultant represents that '
        'it possesses the skills, qualifications, and experience necessary to perform such '
        'services in a competent and professional manner.'
    )

    doc.add_paragraph(
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth '
        'herein, and for other good and valuable consideration, the receipt and sufficiency '
        'of which are hereby acknowledged, the parties agree as follows:'
    )

    # --- Section 1: Scope of Services ---
    doc.add_heading('1. SCOPE OF SERVICES', level=1)

    doc.add_paragraph(
        'The Consultant shall provide the Client with strategic technology consulting services, '
        'including but not limited to: systems architecture review, cloud infrastructure '
        'assessment, cybersecurity audit, and digital transformation roadmap development '
        '(collectively, the "Services"). The specific deliverables and timelines for each '
        'engagement shall be set forth in one or more Statements of Work ("SOW") to be '
        'mutually agreed upon by the parties.'
    )

    doc.add_paragraph(
        'Each Statement of Work shall describe the specific services to be performed, the '
        'deliverables to be provided, the timeline for completion, the fees applicable to '
        'such services, and any special terms or conditions. In the event of any conflict '
        'between the terms of this Agreement and any SOW, the terms of this Agreement shall '
        'prevail unless the SOW expressly states otherwise.'
    )

    # --- Section 2: Compensation ---
    doc.add_heading('2. COMPENSATION AND PAYMENT TERMS', level=1)

    doc.add_paragraph(
        'In consideration of the Services to be performed by the Consultant, the Client '
        'shall pay the Consultant at the rate of $275.00 per hour for senior consultant '
        'services and $175.00 per hour for associate consultant services. All fees shall '
        'be invoiced monthly in arrears, with payment due within thirty (30) calendar days '
        'of the date of invoice.'
    )

    doc.add_paragraph(
        'The Consultant shall be entitled to reimbursement for all reasonable and necessary '
        'out-of-pocket expenses incurred in connection with the performance of the Services, '
        'provided that any single expense exceeding $500.00 shall require prior written '
        'approval from the Client. Expense reports shall be submitted monthly with appropriate '
        'supporting documentation.'
    )

    doc.add_paragraph(
        'Late payments shall accrue interest at the rate of one and one-half percent (1.5%) '
        'per month, or the maximum rate permitted by applicable law, whichever is less, '
        'calculated from the due date until the date of actual payment.'
    )

    # --- Section 3: Term and Termination ---
    doc.add_heading('3. TERM AND TERMINATION', level=1)

    doc.add_paragraph(
        'This Agreement shall commence on the Effective Date and shall continue for a period '
        'of twenty-four (24) months, unless earlier terminated in accordance with the provisions '
        'of this Section. Upon expiration of the initial term, this Agreement shall automatically '
        'renew for successive twelve-month periods unless either party provides written notice '
        'of non-renewal at least ninety (90) days prior to the expiration of the then-current term.'
    )

    doc.add_paragraph(
        'Either party may terminate this Agreement for cause upon thirty (30) days written notice '
        'to the other party if the other party materially breaches any provision of this Agreement '
        'and fails to cure such breach within the notice period. The Client may terminate this '
        'Agreement for convenience upon sixty (60) days written notice to the Consultant, subject '
        'to payment of all fees earned through the effective date of termination.'
    )

    # --- Section 4: Confidentiality ---
    doc.add_heading('4. CONFIDENTIALITY', level=1)

    doc.add_paragraph(
        'Each party acknowledges that in the course of performing its obligations under this '
        'Agreement, it may receive or have access to Confidential Information of the other party. '
        '"Confidential Information" means all non-public information disclosed by one party to '
        'the other, whether orally, in writing, or by inspection, including but not limited to '
        'trade secrets, business plans, financial data, customer lists, technical specifications, '
        'and proprietary methodologies.'
    )

    doc.add_paragraph(
        'The receiving party shall hold all Confidential Information in strict confidence and '
        'shall not disclose such information to any third party without the prior written consent '
        'of the disclosing party. The obligations of confidentiality set forth in this Section '
        'shall survive the termination or expiration of this Agreement for a period of five (5) '
        'years from the date of disclosure.'
    )

    # --- Section 5: Intellectual Property ---
    doc.add_heading('5. INTELLECTUAL PROPERTY RIGHTS', level=1)

    doc.add_paragraph(
        'All work product, deliverables, inventions, and materials created by the Consultant '
        'in the performance of the Services ("Work Product") shall be the sole and exclusive '
        'property of the Client upon full payment of all applicable fees. The Consultant hereby '
        'assigns to the Client all right, title, and interest in and to the Work Product, '
        'including all intellectual property rights therein.'
    )

    doc.add_paragraph(
        'Notwithstanding the foregoing, the Consultant shall retain ownership of all pre-existing '
        'intellectual property, tools, frameworks, and methodologies that were developed by the '
        'Consultant prior to or independently of this Agreement ("Consultant IP"). To the extent '
        'that any Consultant IP is incorporated into the Work Product, the Consultant hereby '
        'grants to the Client a non-exclusive, perpetual, royalty-free license to use such '
        'Consultant IP solely as part of the Work Product.'
    )

    # --- Section 6: Governing Law ---
    doc.add_heading('6. GOVERNING LAW AND DISPUTE RESOLUTION', level=1)

    doc.add_paragraph(
        'This Agreement shall be governed by and construed in accordance with the laws of the '
        'State of Illinois, without regard to its conflict of laws principles. Any dispute arising '
        'out of or relating to this Agreement shall first be submitted to mediation in accordance '
        'with the Commercial Mediation Procedures of the American Arbitration Association. If '
        'mediation is unsuccessful, the dispute shall be resolved by binding arbitration conducted '
        'in Chicago, Illinois.'
    )

    doc.add_paragraph(
        'IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first '
        'written above.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
