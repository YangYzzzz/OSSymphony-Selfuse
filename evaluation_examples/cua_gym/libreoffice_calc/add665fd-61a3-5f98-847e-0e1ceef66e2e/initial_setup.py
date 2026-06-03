"""
Initial Setup: Container management guide in LibreOffice Writer with Docker mock
Task ID: osworld_multi_apps_terminal_screenshot_009
Domain: multi_apps (LibreOffice Writer + Terminal + Screenshot)

Sets up:
1. A container management guide document (DOCX) open in LibreOffice Writer
2. A mock 'docker' command so the agent can run 'docker images' in terminal
3. Opens LibreOffice Writer with the guide
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_terminal_screenshot_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DESKTOP = f'{WORKDIR}/Desktop'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_docker_mock():
    """Create a mock docker command that returns realistic 'docker images' output."""
    mock_script = r"""#!/bin/bash
# Mock docker command for demonstration purposes
if [ "$1" = "images" ]; then
    echo "REPOSITORY                    TAG       IMAGE ID       CREATED        SIZE"
    echo "nginx                         latest    61395b4c586d   2 weeks ago    187MB"
    echo "ubuntu                        22.04     a6be2b5a81a3   3 weeks ago    77.9MB"
    echo "python                        3.11      02e7e7e60035   4 weeks ago    1.01GB"
    echo "redis                         7.2       f9f8e1e3b321   5 weeks ago    138MB"
    echo "postgres                      15        4f8f8e1e3c42   6 weeks ago    412MB"
    echo "node                          18-alpine d3a9e8f4c521   7 weeks ago    127MB"
    echo "alpine                        3.18      8ca4688f4f35   8 weeks ago    7.34MB"
elif [ "$1" = "ps" ]; then
    echo "CONTAINER ID   IMAGE     COMMAND   CREATED   STATUS    PORTS     NAMES"
elif [ "$1" = "--version" ] || [ "$1" = "version" ]; then
    echo "Docker version 24.0.7, build afdd53b"
else
    echo "Usage:  docker [OPTIONS] COMMAND"
    echo ""
    echo "A self-sufficient runtime for containers"
fi
"""
    # Write to /home/user/.local/bin/docker so it's accessible without root
    local_bin = '/home/user/.local/bin'
    os.makedirs(local_bin, exist_ok=True)
    mock_path = f'{local_bin}/docker'
    with open(mock_path, 'w') as f:
        f.write(mock_script)
    os.chmod(mock_path, 0o755)
    print(f'Mock docker command created at {mock_path}')

    # Also add to ~/.bashrc so terminal sessions pick it up
    bashrc_path = '/home/user/.bashrc'
    path_line = f'\nexport PATH="$HOME/.local/bin:$PATH"\n'
    try:
        with open(bashrc_path, 'r') as f:
            content = f.read()
        if '.local/bin' not in content:
            with open(bashrc_path, 'a') as f:
                f.write(path_line)
            print(f'Updated ~/.bashrc to include ~/.local/bin in PATH')
    except FileNotFoundError:
        with open(bashrc_path, 'w') as f:
            f.write(path_line)
        print(f'Created ~/.bashrc with PATH update')


def create_container_guide():
    """Create a container management guide document using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Title
    title = doc.add_heading('Container Management Guide', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction
    doc.add_heading('Introduction', level=1)
    intro = doc.add_paragraph(
        'This guide covers essential Docker commands for managing containers and images '
        'in a development environment. Docker provides a streamlined way to build, ship, '
        'and run applications using containerization technology.'
    )

    # Prerequisites
    doc.add_heading('Prerequisites', level=1)
    prereqs_intro = doc.add_paragraph('Before getting started, ensure you have:')

    # Bullet list
    items = [
        'Docker Engine 24.0 or higher installed',
        'Docker CLI tools available in your PATH',
        'Sufficient disk space for container images (recommended: 20GB+)',
        'Basic familiarity with command-line interfaces',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Section: Viewing Images
    doc.add_heading('Working with Docker Images', level=1)

    doc.add_heading('Listing Available Images', level=2)
    doc.add_paragraph(
        'To view all Docker images currently available on your system, use the following command:'
    )

    # Code block style paragraph
    code_para = doc.add_paragraph()
    run = code_para.add_run('docker images')
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    run.bold = True

    doc.add_paragraph(
        'This command displays a table with the following columns: REPOSITORY, TAG, '
        'IMAGE ID, CREATED, and SIZE. The output helps you identify which images are '
        'available locally and their storage footprint.'
    )

    # Section: Managing Containers
    doc.add_heading('Managing Containers', level=1)

    doc.add_heading('Starting a Container', level=2)
    doc.add_paragraph(
        'To start a new container from an existing image, use the docker run command:'
    )
    run_para = doc.add_paragraph()
    run = run_para.add_run('docker run -d --name my-container nginx:latest')
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    run.bold = True

    doc.add_heading('Listing Running Containers', level=2)
    doc.add_paragraph(
        'To see all currently running containers:'
    )
    ps_para = doc.add_paragraph()
    run = ps_para.add_run('docker ps')
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    run.bold = True

    doc.add_heading('Stopping a Container', level=2)
    doc.add_paragraph(
        'To stop a running container gracefully, use its container ID or name:'
    )
    stop_para = doc.add_paragraph()
    run = stop_para.add_run('docker stop my-container')
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    run.bold = True

    # Section: Image Management
    doc.add_heading('Image Management', level=1)

    doc.add_heading('Pulling Images', level=2)
    doc.add_paragraph(
        'Download images from Docker Hub or other registries using:'
    )
    pull_para = doc.add_paragraph()
    run = pull_para.add_run('docker pull ubuntu:22.04')
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    run.bold = True

    doc.add_heading('Removing Images', level=2)
    doc.add_paragraph(
        'To free up disk space by removing unused images:'
    )
    rmi_para = doc.add_paragraph()
    run = rmi_para.add_run('docker rmi image-id')
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    run.bold = True

    doc.add_paragraph(
        'Note: You cannot remove an image that is currently used by a running container. '
        'Stop and remove the container first using docker stop and docker rm commands.'
    )

    # Best Practices section
    doc.add_heading('Best Practices', level=1)
    doc.add_paragraph(
        'Follow these guidelines to maintain a clean and efficient Docker environment:'
    )

    best_practices = [
        'Regularly audit your images with docker images to track storage usage',
        'Use specific tags (e.g., nginx:1.25) instead of latest for reproducibility',
        'Remove unused images with docker image prune to reclaim disk space',
        'Document your container configurations in docker-compose.yml files',
        'Use multi-stage builds to minimize final image size',
    ]
    for bp in best_practices:
        doc.add_paragraph(bp, style='List Bullet')

    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'This guide provides a foundation for working with Docker containers. '
        'For more advanced topics including networking, volumes, and Docker Compose, '
        'refer to the official Docker documentation at https://docs.docker.com.'
    )

    doc.save(OUTPUT)
    print(f'Container management guide created: {OUTPUT}')


def setup_desktop():
    """Ensure Desktop directory exists."""
    os.makedirs(DESKTOP, exist_ok=True)
    print(f'Desktop directory ready: {DESKTOP}')


def main():
    # 1. Ensure desktop exists
    setup_desktop()

    # 2. Create mock docker command
    create_docker_mock()

    # 3. Create container management guide
    create_container_guide()

    # 4. Launch LibreOffice Writer with the guide
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: LibreOffice Writer launched with container management guide (DISPLAY=:0)')


main()
