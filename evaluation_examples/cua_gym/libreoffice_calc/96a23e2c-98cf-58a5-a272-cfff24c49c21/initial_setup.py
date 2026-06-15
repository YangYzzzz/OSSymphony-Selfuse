"""
Initial Setup: Cookbook tracking spreadsheet and empty document for fewest recipes task
Task ID: osworld_multi_apps_book_reading_rate_010
Domain: libreoffice_calc (multi-app: also involves LibreOffice Writer)
"""

import os
import shlex
import subprocess
import time
import openpyxl
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_book_reading_rate_010'
DESKTOP = '/home/user/Desktop'
XLSX_OUTPUT = f'{DESKTOP}/cookbooks_2022.xlsx'
DOCX_OUTPUT = f'{DESKTOP}/fewest_recipes.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    # --- Create cookbooks_2022.xlsx ---
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Cookbooks 2022"

    # Headers
    headers = ['Title', 'Author', 'Recipe Count']
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)

    # Cookbook data (Title and Author filled, Recipe Count intentionally empty)
    cookbooks = [
        ('Jerusalem', 'Yotam Ottolenghi'),
        ('Salt Fat Acid Heat', 'Samin Nosrat'),
        ('The Food Lab', 'J. Kenji Lopez-Alt'),
        ('Plenty', 'Yotam Ottolenghi'),
        ('Mastering the Art of French Cooking', 'Julia Child'),
    ]

    for row_idx, (title, author) in enumerate(cookbooks, 2):
        ws.cell(row=row_idx, column=1, value=title)
        ws.cell(row=row_idx, column=2, value=author)
        # Recipe Count column is intentionally left empty

    # Adjust column widths for readability
    ws.column_dimensions['A'].width = 45
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 15

    wb.save(XLSX_OUTPUT)
    print(f'Initial spreadsheet created: {XLSX_OUTPUT}')

    # --- Create empty fewest_recipes.docx ---
    doc = Document()
    # Remove default empty paragraph content — leave document truly empty
    # (the default Document() has one blank paragraph, which is fine)
    # Clear out the default paragraph text to keep it empty
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = ''
    doc.save(DOCX_OUTPUT)
    print(f'Empty document created: {DOCX_OUTPUT}')

    # --- GUI-ready startup: open both files ---
    # Open the spreadsheet in LibreOffice Calc
    launch_gui(f'libreoffice --calc "{XLSX_OUTPUT}"', delay_sec=2.0)
    # Open the empty document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DOCX_OUTPUT}"', delay_sec=2.0)

    print('GUI_READY: launched LibreOffice Calc and Writer with DISPLAY=:0')


create_initial()
