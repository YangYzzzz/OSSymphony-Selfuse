"""
Initial Setup: Financial model with audit checklist document
Task ID: osworld_multi_apps_docx_to_calc_011
Domain: libreoffice_calc + libreoffice_writer (multi-app)

Creates:
  - /home/user/financial_model.xlsx  (5 sheets: Inputs, Revenue, Costs, EBITDA, Cashflow)
  - ~/Desktop/audit_checklist.docx   (checklist document describing required validations)
Both files opened in their respective apps.
"""

import os
import shlex
import subprocess
import time
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
DESKTOP = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_docx_to_calc_011'
XLSX_OUTPUT = f'{WORKDIR}/financial_model.xlsx'
DOCX_OUTPUT = f'{DESKTOP}/audit_checklist.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_financial_model():
    """Create financial_model.xlsx with 5 sheets, no validations/formulas added yet."""
    wb = openpyxl.Workbook()

    # ---- Helper styles ----
    header_font = Font(name='Calibri', size=11, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='FF2E4A87', end_color='FF2E4A87', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center')
    thin = Side(style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def style_header(cell):
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = border

    def style_data(cell):
        cell.border = border
        cell.alignment = Alignment(horizontal='left', vertical='center')

    # ========== Sheet 1: Inputs ==========
    ws_inputs = wb.active
    ws_inputs.title = 'Inputs'

    # Section title
    ws_inputs['A1'] = 'Financial Model – Input Parameters'
    ws_inputs['A1'].font = Font(name='Calibri', size=14, bold=True, color='2E4A87')
    ws_inputs.merge_cells('A1:D1')
    ws_inputs['A1'].alignment = Alignment(horizontal='center')

    # Scenario row (NO dropdown — task is to add it)
    ws_inputs['A3'] = 'Scenario'
    ws_inputs['A3'].font = Font(bold=True)
    ws_inputs['B3'] = 'Base Case'   # manually typed, no validation yet

    # Parameter table
    param_headers = ['Parameter', 'Base Case', 'Upside', 'Downside']
    for col, h in enumerate(param_headers, 1):
        cell = ws_inputs.cell(row=5, column=col, value=h)
        style_header(cell)

    params = [
        ('Revenue Growth Rate',    '8%',  '12%',  '4%'),
        ('Cost Growth Rate',       '5%',  '3%',   '7%'),
        ('EBITDA Margin Target',   '22%', '26%',  '18%'),
        ('Capital Expenditure',    '4%',  '3%',   '5%'),
        ('Tax Rate',               '21%', '21%',  '21%'),
        ('Discount Rate (WACC)',   '9%',  '8%',   '10%'),
        ('Working Capital Days',   '45',  '38',   '52'),
        ('Depreciation Rate',      '15%', '15%',  '15%'),
        ('Interest Rate (Debt)',   '5.5%','5.0%', '6.0%'),
        ('Inflation Assumption',   '2.5%','2.0%', '3.5%'),
    ]
    for r, (name, base, up, down) in enumerate(params, 6):
        ws_inputs.cell(row=r, column=1, value=name).border = border
        ws_inputs.cell(row=r, column=2, value=base).border = border
        ws_inputs.cell(row=r, column=3, value=up).border = border
        ws_inputs.cell(row=r, column=4, value=down).border = border

    ws_inputs.column_dimensions['A'].width = 28
    ws_inputs.column_dimensions['B'].width = 14
    ws_inputs.column_dimensions['C'].width = 14
    ws_inputs.column_dimensions['D'].width = 14

    # Notes section
    ws_inputs['A17'] = 'Notes'
    ws_inputs['A17'].font = Font(bold=True, size=11)
    ws_inputs['A18'] = 'All rates are annual. Values represent percentage of revenue unless stated.'
    ws_inputs['A19'] = 'Working Capital Days expressed as days outstanding.'
    ws_inputs['A20'] = 'Scenario selection will drive formula calculations across model sheets.'

    # ========== Sheet 2: Revenue ==========
    ws_rev = wb.create_sheet('Revenue')

    ws_rev['A1'] = 'Revenue Projections ($000)'
    ws_rev['A1'].font = Font(name='Calibri', size=13, bold=True, color='1F6B31')
    ws_rev.merge_cells('A1:G1')
    ws_rev['A1'].alignment = Alignment(horizontal='center')

    rev_headers = ['Product / Segment', 'FY2021A', 'FY2022A', 'FY2023A', 'FY2024E', 'FY2025E', 'FY2026E']
    for col, h in enumerate(rev_headers, 1):
        cell = ws_rev.cell(row=3, column=col, value=h)
        style_header(cell)

    revenue_data = [
        ('Enterprise Software Licenses', 12450, 13680, 14820, 16005, 17285, 18668),
        ('SaaS Subscriptions',           8320,  9875,  11640, 13668, 16054, 18864),
        ('Professional Services',        5230,  5680,  6050,  6534,  7057,  7622),
        ('Maintenance & Support',        4180,  4380,  4620,  4851,  5094,  5348),
        ('Hardware Components',          2640,  2820,  2910,  2968,  3026,  3087),
        ('Training & Certification',     980,   1120,  1260,  1386,  1525,  1677),
        ('Government Contracts',         3450,  3780,  4120,  4532,  4985,  5484),
        ('International Segment',        1890,  2340,  2980,  3726,  4658,  5822),
    ]
    for r, (seg, *vals) in enumerate(revenue_data, 4):
        ws_rev.cell(row=r, column=1, value=seg).border = border
        for c, v in enumerate(vals, 2):
            ws_rev.cell(row=r, column=c, value=v).border = border

    # Total row
    ws_rev.cell(row=12, column=1, value='Total Revenue').font = Font(bold=True)
    ws_rev.cell(row=12, column=1).border = border
    totals = [39140, 43675, 48400, 53670, 59684, 66572]
    for c, v in enumerate(totals, 2):
        cell = ws_rev.cell(row=12, column=c, value=v)
        cell.font = Font(bold=True)
        cell.border = border

    for col_letter in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        ws_rev.column_dimensions[col_letter].width = 24 if col_letter == 'A' else 12

    ws_rev['A14'] = 'YoY Growth'
    ws_rev['A14'].font = Font(bold=True)
    yoy = ['', '', '10.8%', '10.9%', '11.2%', '11.4%']
    for c, v in enumerate(yoy, 2):
        ws_rev.cell(row=14, column=c, value=v)

    # ========== Sheet 3: Costs ==========
    ws_costs = wb.create_sheet('Costs')

    ws_costs['A1'] = 'Cost Structure ($000)'
    ws_costs['A1'].font = Font(name='Calibri', size=13, bold=True, color='8B1A1A')
    ws_costs.merge_cells('A1:G1')
    ws_costs['A1'].alignment = Alignment(horizontal='center')

    cost_headers = ['Cost Category', 'FY2021A', 'FY2022A', 'FY2023A', 'FY2024E', 'FY2025E', 'FY2026E']
    for col, h in enumerate(cost_headers, 1):
        cell = ws_costs.cell(row=3, column=col, value=h)
        style_header(cell)

    costs_data = [
        ('Cost of Goods Sold',           14320, 15980, 17650, 19415, 21584, 24008),
        ('Research & Development',        5480,  6120,  6890,  7579,  8337,  9170),
        ('Sales & Marketing',             6230,  7020,  7980,  8778,  9656,  10621),
        ('General & Administrative',      3780,  4020,  4310,  4569,  4843,  5133),
        ('Depreciation & Amortization',   2140,  2340,  2560,  2742,  2942,  3159),
        ('Restructuring Charges',          320,   180,   240,   120,    80,     50),
        ('Legal & Compliance',             890,   960,  1020,  1071,  1125,  1181),
        ('IT Infrastructure',             1560,  1720,  1920,  2016,  2117,  2223),
        ('Facilities & Lease',            1840,  1920,  2010,  2091,  2175,  2262),
    ]
    for r, (cat, *vals) in enumerate(costs_data, 4):
        ws_costs.cell(row=r, column=1, value=cat).border = border
        for c, v in enumerate(vals, 2):
            ws_costs.cell(row=r, column=c, value=v).border = border

    # Total costs
    ws_costs.cell(row=13, column=1, value='Total Costs').font = Font(bold=True)
    ws_costs.cell(row=13, column=1).border = border
    total_costs = [36560, 40260, 44580, 48381, 52859, 57807]
    for c, v in enumerate(total_costs, 2):
        cell = ws_costs.cell(row=13, column=c, value=v)
        cell.font = Font(bold=True)
        cell.border = border

    for col_letter in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        ws_costs.column_dimensions[col_letter].width = 28 if col_letter == 'A' else 12

    # Manual EBITDA memo (raw values — no cross-sheet formula yet)
    ws_costs['A15'] = 'Memo: EBITDA (manual)'
    ws_costs['A15'].font = Font(bold=True, italic=True)
    memo_ebitda = [2580, 3415, 3820, 5289, 6825, 8765]
    for c, v in enumerate(memo_ebitda, 2):
        ws_costs.cell(row=15, column=c, value=v).font = Font(italic=True)

    # ========== Sheet 4: EBITDA ==========
    ws_ebitda = wb.create_sheet('EBITDA')

    ws_ebitda['A1'] = 'EBITDA Analysis ($000)'
    ws_ebitda['A1'].font = Font(name='Calibri', size=13, bold=True, color='5C3A8A')
    ws_ebitda.merge_cells('A1:G1')
    ws_ebitda['A1'].alignment = Alignment(horizontal='center')

    ebitda_headers = ['Line Item', 'FY2021A', 'FY2022A', 'FY2023A', 'FY2024E', 'FY2025E', 'FY2026E']
    for col, h in enumerate(ebitda_headers, 1):
        cell = ws_ebitda.cell(row=3, column=col, value=h)
        style_header(cell)

    # Revenue and Costs as reference rows (manual values, NOT cross-sheet formulas yet)
    ws_ebitda.cell(row=4, column=1, value='Total Revenue (ref)').border = border
    rev_refs = [39140, 43675, 48400, 53670, 59684, 66572]
    for c, v in enumerate(rev_refs, 2):
        ws_ebitda.cell(row=4, column=c, value=v).border = border

    ws_ebitda.cell(row=5, column=1, value='Total Costs (ref)').border = border
    cost_refs = [36560, 40260, 44580, 48381, 52859, 57807]
    for c, v in enumerate(cost_refs, 2):
        ws_ebitda.cell(row=5, column=c, value=v).border = border

    # EBITDA row — manual values (no cross-sheet formula yet)
    ws_ebitda.cell(row=6, column=1, value='EBITDA').font = Font(bold=True)
    ws_ebitda.cell(row=6, column=1).border = border
    ebitda_vals = [2580, 3415, 3820, 5289, 6825, 8765]
    for c, v in enumerate(ebitda_vals, 2):
        cell = ws_ebitda.cell(row=6, column=c, value=v)
        cell.font = Font(bold=True)
        cell.border = border

    # EBITDA margin (raw percentage values)
    ws_ebitda.cell(row=7, column=1, value='EBITDA Margin %').border = border
    margins = ['6.6%', '7.8%', '7.9%', '9.9%', '11.4%', '13.2%']
    for c, v in enumerate(margins, 2):
        ws_ebitda.cell(row=7, column=c, value=v).border = border

    # D&A add-back
    ws_ebitda.cell(row=9, column=1, value='Add: Depreciation & Amortization').border = border
    ws_ebitda['A9'].font = Font(italic=True)
    da_vals = [2140, 2340, 2560, 2742, 2942, 3159]
    for c, v in enumerate(da_vals, 2):
        ws_ebitda.cell(row=9, column=c, value=v).border = border

    ws_ebitda.cell(row=10, column=1, value='Adjusted EBITDA').font = Font(bold=True)
    ws_ebitda.cell(row=10, column=1).border = border
    adj_ebitda = [4720, 5755, 6380, 8031, 9767, 11924]
    for c, v in enumerate(adj_ebitda, 2):
        cell = ws_ebitda.cell(row=10, column=c, value=v)
        cell.font = Font(bold=True)
        cell.border = border

    ws_ebitda['A12'] = 'NOTE: EBITDA formula (= Revenue - Costs) to be linked via cross-sheet formula audit.'
    ws_ebitda['A12'].font = Font(italic=True, color='808080')

    for col_letter in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        ws_ebitda.column_dimensions[col_letter].width = 30 if col_letter == 'A' else 12

    # ========== Sheet 5: Cashflow ==========
    ws_cf = wb.create_sheet('Cashflow')

    ws_cf['A1'] = 'Cash Flow Statement ($000)'
    ws_cf['A1'].font = Font(name='Calibri', size=13, bold=True, color='1A4A6B')
    ws_cf.merge_cells('A1:G1')
    ws_cf['A1'].alignment = Alignment(horizontal='center')

    cf_headers = ['Cash Flow Item', 'FY2021A', 'FY2022A', 'FY2023A', 'FY2024E', 'FY2025E', 'FY2026E']
    for col, h in enumerate(cf_headers, 1):
        cell = ws_cf.cell(row=3, column=col, value=h)
        style_header(cell)

    cf_data = [
        ('Net Income',                    1840, 2580, 2920, 4012, 5234, 6823),
        ('Add: D&A',                       2140, 2340, 2560, 2742, 2942, 3159),
        ('Change in Working Capital',     -340, -580, -720, -890, -1024, -1187),
        ('Capital Expenditures',         -1560, -1740, -1935, -2147, -2387, -2650),
        ('Free Cash Flow',                2080, 2600, 2825, 3717, 4765, 6145),
        ('', '', '', '', '', '', ''),
        ('Debt Repayment',               -1200, -1200, -1200, -1400, -1400, -1400),
        ('Equity Issuance / (Buyback)',    500,  -200, -300,   0,  -500, -800),
        ('Net Change in Cash',            1380, 1200, 1325, 2317, 2865, 3945),
        ('', '', '', '', '', '', ''),
        ('Opening Cash Balance',         5420, 6800, 8000, 9325, 11642, 14507),
        ('Closing Cash Balance',         6800, 8000, 9325, 11642, 14507, 18452),
    ]
    for r, row_data in enumerate(cf_data, 4):
        for c, v in enumerate(row_data, 1):
            cell = ws_cf.cell(row=r, column=c, value=v)
            if c == 1 and v in ('Free Cash Flow', 'Closing Cash Balance'):
                cell.font = Font(bold=True)
            if v != '':
                cell.border = border

    for col_letter in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        ws_cf.column_dimensions[col_letter].width = 26 if col_letter == 'A' else 12

    wb.save(XLSX_OUTPUT)
    print(f'Initial file created: {XLSX_OUTPUT}')


def create_audit_checklist():
    """Create audit_checklist.docx on Desktop."""
    os.makedirs(DESKTOP, exist_ok=True)
    doc = Document()

    # Title
    title = doc.add_heading('Financial Model Audit Checklist', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    sub = doc.add_paragraph()
    sub_run = sub.add_run('Internal Audit & Validation Requirements — FY2024 Model Review')
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    # Section 1: Overview
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        'This checklist details all required data validations, cross-sheet formulas, and conditional '
        'formatting that must be implemented in financial_model.xlsx to meet audit standards. '
        'Each item below must be fully implemented before the model is submitted for review.'
    )

    # Section 2: Input Validation
    doc.add_heading('2. Input Validation Requirements', level=1)
    doc.add_paragraph('Implement the following data validation rules in the Inputs sheet:')

    items_input = [
        'Add a dropdown list validation to cell B3 (Scenario) with options: "Base Case", "Upside", "Downside". '
        'This allows users to select the active scenario and drives all dependent calculations.',
        'Show the dropdown arrow (set showDropDown to False in DataValidation).',
        'Set prompt title to "Scenario Selection" and prompt message to "Choose the active financial scenario."',
        'Set error title to "Invalid Scenario" and error message to "Please select from: Base Case, Upside, Downside."',
    ]
    for item in items_input:
        doc.add_paragraph(item, style='List Bullet')

    # Section 3: Cross-sheet Formula Checks
    doc.add_heading('3. Cross-Sheet Formula Requirements', level=1)
    doc.add_paragraph('Implement the following cross-sheet linkage formulas:')

    items_formulas = [
        'In the EBITDA sheet, replace all manual "Total Revenue (ref)" values in row 4 '
        '(columns B through G) with cross-sheet formulas referencing Revenue!B12:G12.',
        'In the EBITDA sheet, replace all manual "Total Costs (ref)" values in row 5 '
        '(columns B through G) with cross-sheet formulas referencing Costs!B13:G13.',
        'In the EBITDA sheet, replace EBITDA row 6 (columns B through G) with formulas '
        'computing EBITDA = Revenue - Costs (i.e., =B4-B5 through =G4-G5).',
    ]
    for item in items_formulas:
        doc.add_paragraph(item, style='List Bullet')

    # Section 4: Error Check Sheet
    doc.add_heading('4. Error Check Sheet Requirements', level=1)
    doc.add_paragraph(
        'Create a new sheet named "Error_Check" at the end of the workbook with the following structure:'
    )

    items_error = [
        'Header row (row 1): "Sheet", "Cell", "Computed Value", "Manual Entry", "Deviation %", "Status"',
        'In rows 2 through 7, add IF formulas that compare EBITDA computed values (EBITDA row 6, columns B-G) '
        'against the manual memo entries in Costs sheet row 15 (columns B-G).',
        'Deviation % formula: =ABS((EBITDA!B6 - Costs!B15) / Costs!B15). Use IFERROR to handle zero.',
        'Status formula: =IF(deviation > 5%, "FLAG", "OK"). Threshold is 5% (0.05).',
        'Populate "Sheet" column with "EBITDA", "Cell" column with the cell reference (e.g., "B6"), '
        '"Computed Value" with =EBITDA!B6, "Manual Entry" with =Costs!B15.',
    ]
    for item in items_error:
        doc.add_paragraph(item, style='List Bullet')

    # Section 5: Conditional Formatting
    doc.add_heading('5. Conditional Formatting Requirements', level=1)
    doc.add_paragraph('Apply conditional formatting as follows:')

    items_cf = [
        'In the Error_Check sheet, apply conditional formatting to the Status column (F2:F7): '
        'if cell value equals "FLAG", apply red fill (ARGB: FFFF0000) to highlight discrepancies.',
        'In the EBITDA sheet, apply conditional formatting to EBITDA row 6 (B6:G6): '
        'if the absolute deviation from Costs memo row exceeds 5%, apply red fill (ARGB: FFFF0000).',
    ]
    for item in items_cf:
        doc.add_paragraph(item, style='List Bullet')

    # Section 6: Completion Criteria
    doc.add_heading('6. Completion Criteria', level=1)
    doc.add_paragraph(
        'The audit is considered complete when all of the following are true:'
    )
    completion = [
        'Inputs!B3 has a dropdown validation with exactly three options: Base Case, Upside, Downside.',
        'EBITDA sheet rows 4 and 5 use cross-sheet formula references to Revenue and Costs sheets.',
        'EBITDA sheet row 6 uses =B4-B5 style subtraction formulas.',
        'A sheet named Error_Check exists with IF formulas in the Status column.',
        'Conditional formatting applied to Error_Check!F2:F7 highlighting FLAG rows in red.',
        'File saved as financial_model.xlsx.',
    ]
    for item in completion:
        doc.add_paragraph(item, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run('— End of Audit Checklist —')
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(DOCX_OUTPUT)
    print(f'Audit checklist created: {DOCX_OUTPUT}')


def main():
    create_financial_model()
    create_audit_checklist()

    # GUI-ready startup: open audit_checklist.docx in Writer first, then financial_model.xlsx in Calc
    launch_gui(f'libreoffice --writer "{DOCX_OUTPUT}"', delay_sec=2.0)
    launch_gui(f'libreoffice --calc "{XLSX_OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer and Calc with DISPLAY=:0')


main()
