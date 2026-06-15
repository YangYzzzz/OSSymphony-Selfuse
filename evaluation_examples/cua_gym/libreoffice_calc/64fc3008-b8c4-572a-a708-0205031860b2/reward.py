"""
Reward Script: Download paper PDF and check citation in docx
Task ID: osworld_multi_apps_pdf_download_cite_001
Domain: multi_apps (PDF download + LibreOffice Writer)
Scoring:
  Component 1: paper01.pdf exists and is a valid PDF with correct content (0.4 pts)
  Component 2: citation_check.docx exists and is a valid docx (0.2 pts)
  Component 3: citation_check.docx contains 'Yes' answer about BERT citing Transformer (0.4 pts)
  Total: 1.0
"""

import os
import re
import zlib

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_pdf_download_cite_001'

PDF_PATH = f'{WORKDIR}/paper01.pdf'
DOCX_PATH = f'{WORKDIR}/citation_check.docx'


def verify_task():
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Component 1: paper01.pdf exists, is a valid PDF, and contains content from
    # "Attention Is All You Need" paper (arXiv:1706.03762) (0.4 points)
    # This should FAIL on initial_env (file does not exist) and PASS on golden_env
    try:
        if not os.path.isfile(PDF_PATH):
            print(f"FAIL: Component 1 — paper01.pdf not found at {PDF_PATH}")
        else:
            file_size = os.path.getsize(PDF_PATH)
            if file_size < 100:
                print(f"FAIL: Component 1 — paper01.pdf too small ({file_size} bytes), likely not a real PDF")
            else:
                with open(PDF_PATH, 'rb') as f:
                    pdf_bytes = f.read()
                # Check PDF magic header
                if not pdf_bytes.startswith(b'%PDF'):
                    print(f"FAIL: Component 1 — paper01.pdf missing PDF header, found: {pdf_bytes[:8]}")
                else:
                    # Decompress PDF streams to find "Attention Is All You Need" title
                    streams = re.findall(rb'stream\r?\n(.+?)endstream', pdf_bytes, re.DOTALL)
                    stream_text_combined = ''
                    for stream in streams:
                        try:
                            decompressed = zlib.decompress(stream.strip())
                            stream_text_combined += decompressed.decode('latin-1', errors='replace')
                        except Exception:
                            # Non-FlateDecode stream, check raw
                            stream_text_combined += stream.decode('latin-1', errors='replace')

                    # Also include raw PDF bytes for metadata/annotation checks
                    raw_text = pdf_bytes.decode('latin-1', errors='replace')
                    all_text = stream_text_combined + raw_text

                    # Check for paper title or ArXiv ID
                    attention_keywords_found = sum([
                        1 if 'Attention Is All You Need' in all_text else 0,
                        1 if '1706.03762' in all_text else 0,
                        1 if 'Vaswani' in all_text else 0,
                    ])

                    if attention_keywords_found >= 1:
                        print(f"PASS: Component 1 — paper01.pdf is valid PDF ({file_size} bytes) "
                              f"containing 'Attention Is All You Need' content ({attention_keywords_found} keywords matched) (0.4 pts)")
                        total_score += 0.4
                    else:
                        print(f"FAIL: Component 1 — paper01.pdf exists ({file_size} bytes) but does not "
                              f"contain expected 'Attention Is All You Need' content (0 keywords matched)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: citation_check.docx exists and is a valid Word document (0.2 points)
    # This should FAIL on initial_env (file does not exist) and PASS on golden_env
    try:
        if not os.path.isfile(DOCX_PATH):
            print(f"FAIL: Component 2 — citation_check.docx not found at {DOCX_PATH}")
        else:
            from docx import Document
            try:
                doc = Document(DOCX_PATH)
                num_paras = len(doc.paragraphs)
                if num_paras > 0:
                    print(f"PASS: Component 2 — citation_check.docx is valid docx with "
                          f"{num_paras} paragraphs (0.2 pts)")
                    total_score += 0.2
                else:
                    print(f"FAIL: Component 2 — citation_check.docx is empty (0 paragraphs)")
            except Exception as e:
                print(f"FAIL: Component 2 — citation_check.docx exists but cannot be opened: {e}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: citation_check.docx contains a 'Yes' answer indicating BERT cites
    # the Transformer paper, with explanation (0.4 points)
    # This should FAIL on initial_env (file does not exist) and PASS on golden_env
    try:
        if not os.path.isfile(DOCX_PATH):
            print(f"FAIL: Component 3 — citation_check.docx not found, cannot check citation answer")
        else:
            from docx import Document
            try:
                doc = Document(DOCX_PATH)
                full_text = '\n'.join(para.text for para in doc.paragraphs)
                full_text_lower = full_text.lower()

                # Must contain "Yes" as the answer (case-insensitive)
                has_yes = bool(re.search(r'\byes\b', full_text_lower))

                # Must contain some explanation referencing BERT or the Transformer
                has_explanation = (
                    'bert' in full_text_lower or
                    'transformer' in full_text_lower or
                    'attention' in full_text_lower or
                    'cite' in full_text_lower or
                    'vaswani' in full_text_lower
                )

                if has_yes and has_explanation:
                    print(f"PASS: Component 3 — citation_check.docx contains 'Yes' answer with "
                          f"explanation referencing the cited paper (0.4 pts)")
                    total_score += 0.4
                elif has_yes and not has_explanation:
                    print(f"FAIL: Component 3 — citation_check.docx contains 'Yes' but lacks "
                          f"explanation about why BERT cites Transformer")
                elif not has_yes:
                    # Check if "No" was incorrectly given
                    has_no = bool(re.search(r'\bno\b', full_text_lower))
                    if has_no:
                        print(f"FAIL: Component 3 — citation_check.docx contains 'No' answer, "
                              f"but expected 'Yes' (BERT does cite Attention Is All You Need)")
                    else:
                        print(f"FAIL: Component 3 — citation_check.docx does not contain 'Yes' "
                              f"answer. Full text snippet: {full_text[:300]}")
                else:
                    print(f"FAIL: Component 3 — citation_check.docx missing required citation answer")
            except Exception as e:
                print(f"FAIL: Component 3 — citation_check.docx cannot be parsed: {e}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


verify_task()
