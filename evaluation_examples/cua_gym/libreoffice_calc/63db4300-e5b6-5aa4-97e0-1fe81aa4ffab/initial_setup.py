"""
Initial Setup: Thesis chapter with incorrectly formatted bibliography
Task ID: osworld_multi_apps_misc_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_misc_043'
DRAFTS_DIR = f'{WORKDIR}/Desktop/drafts'
OUTPUT = f'{DRAFTS_DIR}/thesis_chapter.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure the drafts directory exists
    os.makedirs(DRAFTS_DIR, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading('Chapter 3: The Role of Digital Archives in Historical Research', level=1)

    # Introduction section
    doc.add_heading('3.1 Introduction', level=2)
    doc.add_paragraph(
        'The proliferation of digital archives over the past two decades has fundamentally '
        'transformed the practice of historical research. Scholars now have access to primary '
        'sources that were previously inaccessible due to geographic or institutional barriers. '
        'This chapter examines the methodological implications of these developments and surveys '
        'existing scholarship on the topic.'
    )

    doc.add_paragraph(
        'Early work by Henderson (2014) established the theoretical groundwork for understanding '
        'how digitization affects source selection and interpretation. Subsequent studies have '
        'expanded this framework to consider issues of preservation, metadata standards, and '
        'long-term accessibility (Morris and Patel 2016; Nakamura 2018).'
    )

    # Section 2
    doc.add_heading('3.2 Methodological Considerations', level=2)
    doc.add_paragraph(
        'When utilizing digital archives, researchers must consider several methodological '
        'challenges. First, digitization projects are selective, meaning that not all documents '
        'within an archive are digitized equally. As Thompson (2015) notes, "the biases inherent '
        'in digitization decisions can skew our understanding of historical periods and populations" (p. 47).'
    )

    doc.add_paragraph(
        'Second, the metadata associated with digitized documents varies considerably across '
        'institutions. Standardization efforts led by organizations such as the Digital Preservation '
        'Coalition have improved consistency, but significant disparities remain (Okafor and Singh 2019). '
        'These disparities complicate cross-institutional comparative research and require scholars '
        'to develop robust data-cleaning protocols.'
    )

    doc.add_paragraph(
        'Third, optical character recognition (OCR) technologies, while increasingly sophisticated, '
        'still introduce transcription errors that can affect textual analysis at scale. Researchers '
        'working with large corpora must implement verification procedures to ensure data integrity '
        '(Yamamoto 2017).'
    )

    # Section 3
    doc.add_heading('3.3 Case Studies', level=2)
    doc.add_paragraph(
        'Several recent projects illustrate both the potential and the limitations of digital '
        'archive research. The Colonial Dispatches project at the University of Victoria demonstrated '
        'how digitized government records could reveal patterns invisible to traditional archival '
        'researchers (Chen and Williams 2020). By analyzing over 40,000 documents spanning '
        'three decades, the team identified previously undocumented communication networks '
        'between colonial administrators.'
    )

    doc.add_paragraph(
        'In contrast, the Digital Scriptorium project encountered significant challenges when '
        'attempting to standardize metadata across participating institutions (Rivera 2021). '
        'Despite years of collaboration, inconsistencies in cataloging practices persisted, '
        'ultimately limiting the project\'s comparative analytical capacity.'
    )

    # Section 4
    doc.add_heading('3.4 Future Directions', level=2)
    doc.add_paragraph(
        'The future of digital archive research lies in the development of more sophisticated '
        'machine learning tools capable of automating metadata extraction and improving OCR '
        'accuracy. Promising work in this area has been conducted by researchers at several '
        'institutions (Kumar and Lee 2022; Bergmann 2023). However, these tools require '
        'significant computational resources and specialized expertise, raising questions '
        'about equity of access within the historical research community.'
    )

    doc.add_paragraph(
        'Additionally, the long-term preservation of digital archives remains a pressing concern. '
        'Format obsolescence, institutional instability, and funding constraints all threaten '
        'the continued accessibility of digitized collections. Addressing these challenges will '
        'require sustained collaboration between librarians, archivists, technologists, and '
        'funding bodies.'
    )

    # Bibliography section (INCORRECTLY FORMATTED - mixed styles, missing dates, wrong punctuation)
    doc.add_heading('Bibliography', level=1)

    doc.add_paragraph(
        'Bergmann, Klaus (2023). Automated Metadata Extraction Using Neural Networks. '
        'Journal of Digital Humanities, 12(3), 88-107.'
    )

    doc.add_paragraph(
        'Chen, Mei and Williams, Robert. 2020. "Colonial Networks in the Digital Age." '
        'Pacific Historical Review 89 (2): 210-245.'
    )

    doc.add_paragraph(
        'Henderson, Patricia. 2014. Digital Archives and Historical Method. '
        'Oxford: Oxford University Press.'
    )

    doc.add_paragraph(
        'Kumar, Arjun and Lee, Soo-Jin (2022). Machine Learning Applications in Archival Research. '
        'Boston: MIT Press.'
    )

    doc.add_paragraph(
        'Morris, James and Patel, Anita, 2016. Preservation Standards in Digital Libraries. '
        'London: Routledge.'
    )

    doc.add_paragraph(
        'Nakamura, Yuki 2018. Metadata Frameworks for Historical Digitization Projects. '
        'New York: Columbia University Press.'
    )

    doc.add_paragraph(
        'Okafor, Chidinma and Singh, Harbhajan. 2019. Cross-Institutional Data Consistency '
        'in Digital Archives. Cambridge: Cambridge University Press.'
    )

    doc.add_paragraph(
        'Rivera, Maria. 2021. "Challenges in Multi-Institutional Cataloging Projects." '
        'Library Quarterly 91(4): 312-330.'
    )

    doc.add_paragraph(
        'Thompson, David. 2015. Bias and Selection in Digital Preservation. '
        'Chicago: University of Chicago Press.'
    )

    doc.add_paragraph(
        'Yamamoto, Hiroshi 2017. OCR Errors and Textual Analysis: A Methodological Study. '
        'Edinburgh: Edinburgh University Press.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the thesis chapter in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
