"""
Initial Setup: report_draft.docx with non-compliant APA 7th edition bibliography
Task ID: osworld_multi_apps_misc_048
Domain: libreoffice_writer

Creates /home/user/Desktop/team_docs/report_draft.docx with:
- A research report document with multiple sections
- A References section with 5 bibliography entries that are NOT APA 7th compliant:
  * Some lack hanging indent formatting
  * Some have Title Case titles (should be sentence case)
  * Some journal articles include city of publication (should be omitted in APA 7)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_misc_048'
DOCS_FOLDER = f'{WORKDIR}/Desktop/team_docs'
OUTPUT = f'{DOCS_FOLDER}/report_draft.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Create the team_docs folder if not exists
    os.makedirs(DOCS_FOLDER, exist_ok=True)

    doc = Document()

    # Set up page margins (APA-style: 1 inch all sides)
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ---- Title ----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("Climate Change Adaptation Strategies in Urban Environments")
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()  # blank line

    # Author
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.add_run("Dr. Alexandra Nguyen, Dr. James Whitfield, & Ms. Priya Kapoor")

    affil_para = doc.add_paragraph()
    affil_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    affil_para.add_run("Department of Environmental Studies, Greenfield University")

    doc.add_paragraph()  # blank line

    # ---- Abstract ----
    abstract_heading = doc.add_paragraph()
    abstract_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    abstract_run = abstract_heading.add_run("Abstract")
    abstract_run.bold = True

    abstract_body = doc.add_paragraph(
        "This paper examines climate change adaptation strategies employed in major urban centers "
        "across North America and Europe between 2010 and 2024. Drawing on longitudinal data from "
        "47 municipalities, we identify five key intervention categories: green infrastructure, "
        "early-warning systems, building code revisions, community resilience programs, and "
        "economic incentive schemes. Our analysis reveals that combined approaches outperform "
        "single-domain strategies by an average of 34% on a standardized resilience index (SRI). "
        "Policy implications for city planners and national governments are discussed."
    )
    abstract_body.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_paragraph()

    # ---- Introduction ----
    intro_heading = doc.add_paragraph()
    intro_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    intro_run = intro_heading.add_run("Introduction")
    intro_run.bold = True

    doc.add_paragraph(
        "Urban areas are responsible for approximately 70% of global greenhouse gas emissions "
        "yet house more than half of the world's population (United Nations, 2023). As climate "
        "change accelerates, cities face mounting pressure to adapt their infrastructure, "
        "governance systems, and community services. The urgency of this challenge is underscored "
        "by increasingly frequent extreme weather events, sea-level rise, and shifting precipitation "
        "patterns (IPCC, 2022)."
    ).paragraph_format.first_line_indent = Inches(0.5)

    doc.add_paragraph(
        "Previous reviews have examined adaptation strategies in isolation—focusing on either "
        "technical solutions or social programs—without assessing their combined effectiveness. "
        "This study addresses that gap by applying a mixed-methods framework to compare integrated "
        "versus siloed approaches across a diverse sample of cities."
    ).paragraph_format.first_line_indent = Inches(0.5)

    # ---- Methods ----
    methods_heading = doc.add_paragraph()
    methods_run = methods_heading.add_run("Methods")
    methods_run.bold = True

    doc.add_paragraph(
        "Data were collected from municipal sustainability reports, academic publications, and "
        "structured interviews with 112 urban planners and policymakers. Quantitative analysis "
        "used a standardized resilience index (SRI) developed by Hartmann and Osei (2021), "
        "validated across 15 climate zones. Qualitative themes were derived through thematic "
        "analysis following Braun and Clarke (2006)."
    ).paragraph_format.first_line_indent = Inches(0.5)

    # ---- Results ----
    results_heading = doc.add_paragraph()
    results_run = results_heading.add_run("Results")
    results_run.bold = True

    doc.add_paragraph(
        "Cities implementing at least three concurrent adaptation strategies scored significantly "
        "higher on the SRI (M = 6.8, SD = 1.2) compared to those using a single strategy "
        "(M = 4.1, SD = 1.5), t(45) = 7.34, p < .001. Green infrastructure investments showed "
        "the strongest independent effect (β = 0.42, SE = 0.09), followed by community resilience "
        "programs (β = 0.31, SE = 0.11)."
    ).paragraph_format.first_line_indent = Inches(0.5)

    # ---- Discussion ----
    discussion_heading = doc.add_paragraph()
    discussion_run = discussion_heading.add_run("Discussion")
    discussion_run.bold = True

    doc.add_paragraph(
        "These findings align with the integrated resilience model proposed by Torres-Ruiz et al. "
        "(2020), suggesting that interdependencies between physical and social systems amplify "
        "adaptive capacity. Policymakers should prioritize cross-sector collaboration and avoid "
        "piecemeal investments that fail to leverage synergistic effects."
    ).paragraph_format.first_line_indent = Inches(0.5)

    doc.add_paragraph(
        "Limitations include the reliance on self-reported municipal data, which may introduce "
        "social desirability bias. Future research should employ independent auditing of adaptation "
        "program outcomes and extend the temporal scope beyond 2024."
    ).paragraph_format.first_line_indent = Inches(0.5)

    # ---- Conclusion ----
    conclusion_heading = doc.add_paragraph()
    conclusion_run = conclusion_heading.add_run("Conclusion")
    conclusion_run.bold = True

    doc.add_paragraph(
        "Urban climate adaptation is most effective when strategies are integrated across "
        "infrastructure, governance, and community domains. This research provides actionable "
        "guidance for city planners and establishes a foundation for longitudinal assessments "
        "of adaptation outcomes."
    ).paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # ---- References heading ----
    ref_heading = doc.add_paragraph()
    ref_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ref_heading_run = ref_heading.add_run("References")
    ref_heading_run.bold = True

    doc.add_paragraph()  # blank line after heading

    # ---- 5 Non-compliant APA references ----
    # Issue 1: Title Case (should be sentence case), no hanging indent, journal article with city
    ref1 = doc.add_paragraph()
    ref1.paragraph_format.left_indent = Inches(0)
    ref1.paragraph_format.first_line_indent = Inches(0)
    ref1.add_run(
        "Braun, V., & Clarke, V. (2006). Using Thematic Analysis in Psychology. "
        "Qualitative Research in Psychology, 3(2), 77–101. London. "
        "https://doi.org/10.1191/1478088706qp063oa"
    )

    # Issue 2: Title Case, no hanging indent
    ref2 = doc.add_paragraph()
    ref2.paragraph_format.left_indent = Inches(0)
    ref2.paragraph_format.first_line_indent = Inches(0)
    ref2.add_run(
        "Hartmann, D., & Osei, K. (2021). A Standardized Resilience Index For Urban Climate "
        "Assessment: Development And Validation Across Fifteen Climate Zones. "
        "Urban Studies, 58(4), 812–831. New York. "
        "https://doi.org/10.1177/0042098020920739"
    )

    # Issue 3: Title Case, no hanging indent, journal article with city
    ref3 = doc.add_paragraph()
    ref3.paragraph_format.left_indent = Inches(0)
    ref3.paragraph_format.first_line_indent = Inches(0)
    ref3.add_run(
        "IPCC. (2022). Climate Change 2022: Impacts, Adaptation And Vulnerability. "
        "Contribution Of Working Group II To The Sixth Assessment Report Of The "
        "Intergovernmental Panel On Climate Change. Cambridge University Press."
    )

    # Issue 4: Title Case in article title, has city of publication for journal
    ref4 = doc.add_paragraph()
    ref4.paragraph_format.left_indent = Inches(0)
    ref4.paragraph_format.first_line_indent = Inches(0)
    ref4.add_run(
        "Torres-Ruiz, A., Morales, C., & Fitzpatrick, E. (2020). Interdependencies In Urban "
        "Resilience: Physical Infrastructure And Social Capital As Complementary Adaptive Systems. "
        "Nature Climate Change, 10(7), 621–628. Berlin. "
        "https://doi.org/10.1038/s41558-020-0789-6"
    )

    # Issue 5: Title Case, no hanging indent, journal article with city
    ref5 = doc.add_paragraph()
    ref5.paragraph_format.left_indent = Inches(0)
    ref5.paragraph_format.first_line_indent = Inches(0)
    ref5.add_run(
        "United Nations. (2023). World Urbanization Prospects: The 2023 Revision. "
        "United Nations Department of Economic and Social Affairs. New York. "
        "https://population.un.org/wup/"
    )

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
