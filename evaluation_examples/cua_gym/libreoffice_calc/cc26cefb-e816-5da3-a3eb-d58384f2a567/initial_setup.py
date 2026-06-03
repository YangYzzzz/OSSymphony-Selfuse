"""
Initial Setup: Grant Proposal with inconsistent Vancouver citations
Task ID: osworld_multi_apps_misc_047
Domain: libreoffice_writer (docx)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_misc_047'
PROPOSALS_DIR = f'{WORKDIR}/Desktop/proposals'
OUTPUT = f'{PROPOSALS_DIR}/grant_proposal.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Create proposals directory on Desktop
    os.makedirs(PROPOSALS_DIR, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading('Research Grant Proposal', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Section: Project Summary
    doc.add_heading('1. Project Summary', level=1)
    doc.add_paragraph(
        'This proposal seeks funding to investigate novel therapeutic approaches for '
        'treatment-resistant cardiovascular disease. Our interdisciplinary team will '
        'combine cutting-edge genomics, clinical trials, and bioinformatics to identify '
        'new biomarkers and develop targeted interventions. The project spans three years '
        'and involves collaborations with leading medical institutions.'
    )

    # Section: Background and Significance
    doc.add_heading('2. Background and Significance', level=1)
    doc.add_paragraph(
        'Cardiovascular disease remains the leading cause of mortality worldwide, '
        'accounting for approximately 17.9 million deaths annually. Despite significant '
        'advances in treatment, a substantial proportion of patients exhibit resistance '
        'to standard therapies. Recent studies have highlighted the role of epigenetic '
        'modifications and microRNA dysregulation in treatment failure [1,2].'
    )
    doc.add_paragraph(
        'The RAAS pathway has been extensively studied in the context of hypertension '
        'and heart failure. However, its interaction with inflammatory mediators remains '
        'incompletely characterized [3]. Furthermore, emerging evidence suggests that '
        'gut microbiome composition influences drug metabolism and treatment outcomes [4,5].'
    )

    # Section: Specific Aims
    doc.add_heading('3. Specific Aims', level=1)
    doc.add_paragraph('Aim 1: To identify epigenetic biomarkers associated with treatment resistance.')
    doc.add_paragraph('Aim 2: To characterize microRNA expression profiles in non-responders.')
    doc.add_paragraph('Aim 3: To evaluate gut microbiome-drug interactions in cardiovascular patients.')
    doc.add_paragraph('Aim 4: To develop a predictive model for treatment response stratification.')

    # Section: Research Strategy
    doc.add_heading('4. Research Strategy', level=1)
    doc.add_paragraph(
        'We will recruit 500 patients with treatment-resistant hypertension from three '
        'academic medical centers. Whole-genome bisulfite sequencing will be performed '
        'to generate comprehensive DNA methylation profiles. MicroRNA sequencing will '
        'be conducted using Illumina NextSeq 500 platforms. Metagenomic analysis will '
        'characterize the gut microbiome composition and predicted functional capacity.'
    )
    doc.add_paragraph(
        'Statistical analyses will employ mixed-effects models to account for '
        'longitudinal data structure and inter-site variability. Machine learning '
        'algorithms, including random forest and gradient boosting, will be used for '
        'biomarker discovery and predictive modeling. All analyses will be pre-registered '
        'on ClinicalTrials.gov prior to study initiation.'
    )

    # Section: Budget Justification
    doc.add_heading('5. Budget Justification', level=1)
    doc.add_paragraph(
        'The requested budget of $1,250,000 over three years covers personnel (60%), '
        'equipment and supplies (25%), and indirect costs (15%). Personnel include two '
        'postdoctoral researchers, one biostatistician, and 0.5 FTE of senior investigator '
        'effort. Major equipment purchases include next-generation sequencing consumables '
        'and bioinformatics cloud computing resources.'
    )

    # Section: References - with intentionally inconsistent Vancouver formatting
    doc.add_heading('6. References', level=1)

    # Reference 1: Missing journal abbreviation (uses full journal name instead)
    doc.add_paragraph(
        '1. Zhang WH, Li M, Thompson RA, Garcia JF, Patel S. '
        'Epigenetic modifications in treatment-resistant hypertension: a genome-wide analysis. '
        'The New England Journal of Medicine. 2023;388(14):1285-1298. '
        'doi:10.1056/NEJMoa2300847. PMID: 37018465.'
    )

    # Reference 2: Volume/issue format wrong (uses Vol. X, No. Y format instead of X(Y))
    doc.add_paragraph(
        '2. Chen L, Nakamura T, Williams SJ, Rodriguez E. '
        'MicroRNA dysregulation and cardiovascular drug resistance: mechanistic insights. '
        'Circulation. Vol. 147, No. 8 (2023):pp. 612-625.'
    )

    # Reference 3: Missing issue number, trailing extra info
    doc.add_paragraph(
        '3. Anderson KP, Brown MR, Liu X, Schmidt H, Fernandez D, Kim JY. '
        'RAAS pathway interactions with inflammatory mediators in heart failure: '
        'a systematic review and meta-analysis. '
        'Journal of the American College of Cardiology. 2022;80:445-462. '
        'Impact factor: 24.0. Highly cited paper.'
    )

    # Reference 4: Incorrect format - page range missing, extra text in wrong position
    doc.add_paragraph(
        '4. Petrov AI, Sanchez-Martinez C, White H, Nguyen TT, Okonkwo B. '
        'Gut microbiome composition predicts metoprolol metabolism in hypertensive patients. '
        'Nature Medicine. 2023, Volume 29, Issue 3, March 2023: pages 678-691. '
        'Epub ahead of print available at: https://doi.org/10.1038/s41591-023-02245-7'
    )

    # Reference 5: Uses journal's full name AND wrong punctuation
    doc.add_paragraph(
        '5. Osei-Bonsu K, Tran QH, Millar DA, Johansson E. '
        'Predictive biomarkers for statin therapy in resistant hyperlipidemia. '
        'European Heart Journal (Oxford University Press). 2022:43(44);4218-4231.'
    )

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Create other documents in the proposals folder for realistic context
    other_doc_names = [
        'budget_justification.docx',
        'team_cv_summary.docx',
        'institutional_support_letter.docx',
    ]
    for fname in other_doc_names:
        other_path = os.path.join(PROPOSALS_DIR, fname)
        if not os.path.exists(other_path):
            odoc = Document()
            odoc.add_heading(fname.replace('_', ' ').replace('.docx', '').title(), level=0)
            odoc.add_paragraph('This document is part of the grant proposal package.')
            odoc.save(other_path)

    # GUI-ready startup: open the grant_proposal.docx in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with grant_proposal.docx (DISPLAY=:0)')


create_initial()
