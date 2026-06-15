"""
Initial Setup: Merge Sort Tutorial and Python Stub
Task ID: osworld_multi_apps_misc_033
Domain: multi_apps (libreoffice_writer + os/python)

Creates:
  - /home/user/Desktop/Merge_Sort_Tutorial.docx  (opened in LibreOffice Writer)
  - /home/user/Desktop/mergeSort.py               (incomplete stub with test cases)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

DESKTOP = '/home/user/Desktop'
TUTORIAL_PATH = f'{DESKTOP}/Merge_Sort_Tutorial.docx'
STUB_PATH = f'{DESKTOP}/mergeSort.py'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_tutorial_docx():
    """Create the Merge Sort Tutorial document."""
    doc = Document()

    # Title
    title = doc.add_heading('Merge Sort Algorithm Tutorial', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'Merge Sort is a classic divide-and-conquer sorting algorithm. '
        'It works by recursively splitting an array into two halves, '
        'sorting each half, and then merging the two sorted halves back together. '
        'Merge Sort has a guaranteed time complexity of O(n log n) for all cases, '
        'making it more reliable than algorithms like Quick Sort for worst-case scenarios.'
    )

    # How It Works
    doc.add_heading('How Merge Sort Works', level=1)
    doc.add_paragraph(
        'Merge Sort follows these key steps:'
    )
    doc.add_paragraph('1. Divide: Split the unsorted list into two approximately equal halves.', style='List Number')
    doc.add_paragraph('2. Conquer: Recursively sort each half using Merge Sort.', style='List Number')
    doc.add_paragraph('3. Combine: Merge the two sorted halves into a single sorted list.', style='List Number')

    doc.add_paragraph(
        'The base case of the recursion is when the list has zero or one element, '
        'as a list with a single element is already sorted.'
    )

    # Pseudocode
    doc.add_heading('Pseudocode', level=1)
    doc.add_paragraph(
        'The following pseudocode describes the Merge Sort algorithm:'
    )

    pseudocode_para = doc.add_paragraph()
    run = pseudocode_para.add_run(
        'function mergeSort(arr):\n'
        '    if length(arr) <= 1:\n'
        '        return arr\n'
        '    mid = length(arr) / 2\n'
        '    left = mergeSort(arr[0 : mid])\n'
        '    right = mergeSort(arr[mid : end])\n'
        '    return merge(left, right)\n'
        '\n'
        'function merge(left, right):\n'
        '    result = []\n'
        '    i = 0, j = 0\n'
        '    while i < length(left) and j < length(right):\n'
        '        if left[i] <= right[j]:\n'
        '            append left[i] to result\n'
        '            i = i + 1\n'
        '        else:\n'
        '            append right[j] to result\n'
        '            j = j + 1\n'
        '    append remaining elements of left to result\n'
        '    append remaining elements of right to result\n'
        '    return result'
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # Example
    doc.add_heading('Step-by-Step Example', level=1)
    doc.add_paragraph(
        'Consider sorting the array: [38, 27, 43, 3, 9, 82, 10]'
    )

    doc.add_paragraph('Step 1 — Divide:')
    doc.add_paragraph('[38, 27, 43, 3, 9, 82, 10]  →  [38, 27, 43]  and  [3, 9, 82, 10]',
                      style='List Bullet')
    doc.add_paragraph('[38, 27, 43]  →  [38, 27]  and  [43]', style='List Bullet')
    doc.add_paragraph('[38, 27]  →  [38]  and  [27]', style='List Bullet')

    doc.add_paragraph('Step 2 — Conquer (merge sorted sub-lists):')
    doc.add_paragraph('merge([38], [27])  →  [27, 38]', style='List Bullet')
    doc.add_paragraph('merge([27, 38], [43])  →  [27, 38, 43]', style='List Bullet')
    doc.add_paragraph('merge([3], [9])  →  [3, 9]', style='List Bullet')
    doc.add_paragraph('merge([82], [10])  →  [10, 82]', style='List Bullet')
    doc.add_paragraph('merge([3, 9], [10, 82])  →  [3, 9, 10, 82]', style='List Bullet')

    doc.add_paragraph('Step 3 — Final Merge:')
    doc.add_paragraph('merge([27, 38, 43], [3, 9, 10, 82])  →  [3, 9, 10, 27, 38, 43, 82]',
                      style='List Bullet')

    doc.add_paragraph(
        'Final sorted array: [3, 9, 10, 27, 38, 43, 82]'
    )

    # Complexity
    doc.add_heading('Time and Space Complexity', level=1)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Case'
    hdr[1].text = 'Time Complexity'
    hdr[2].text = 'Space Complexity'
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    rows_data = [
        ('Best Case',    'O(n log n)', 'O(n)'),
        ('Average Case', 'O(n log n)', 'O(n)'),
        ('Worst Case',   'O(n log n)', 'O(n)'),
    ]
    for i, (case, time_c, space_c) in enumerate(rows_data, 1):
        row = table.rows[i].cells
        row[0].text = case
        row[1].text = time_c
        row[2].text = space_c

    doc.add_paragraph(
        '\nMerge Sort requires O(n) auxiliary space for the temporary arrays used during merging. '
        'This is a trade-off compared to in-place algorithms, but the guaranteed O(n log n) '
        'time makes it preferable for large datasets.'
    )

    # Implementation Hints
    doc.add_heading('Implementation Hints', level=1)
    doc.add_paragraph(
        'When implementing mergeSort in Python, remember the following:'
    )
    doc.add_paragraph(
        'The function should accept a list and return a new sorted list.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Use Python list slicing (arr[:mid] and arr[mid:]) to split the array.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Create a helper merge() function or implement inline merging using two pointers.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'The base case: if len(arr) <= 1, return arr immediately.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'After comparing elements, append any remaining elements from left or right sub-array.',
        style='List Bullet'
    )

    os.makedirs(DESKTOP, exist_ok=True)
    doc.save(TUTORIAL_PATH)
    print(f'Tutorial created: {TUTORIAL_PATH}')


def create_stub_py():
    """Create the incomplete mergeSort.py stub with test cases."""
    stub_code = '''\
#!/usr/bin/env python3
"""
mergeSort.py - Merge Sort Implementation Exercise
Complete the mergeSort function below based on the tutorial in Merge_Sort_Tutorial.docx
"""


def mergeSort(arr):
    """
    Sort a list using the merge sort algorithm.

    Args:
        arr (list): The input list to be sorted.

    Returns:
        list: A new sorted list containing all elements from arr.

    TODO: Implement this function.
          Hint: Base case - if len(arr) <= 1, return arr.
          Hint: Find the midpoint, recursively sort left and right halves,
                then merge the two sorted halves together.
    """
    # TODO: implement merge sort here
    pass


# ── Test Cases ──────────────────────────────────────────────────────────────

def run_tests():
    test_cases = [
        {
            "description": "Example from tutorial",
            "input": [38, 27, 43, 3, 9, 82, 10],
            "expected": [3, 9, 10, 27, 38, 43, 82],
        },
        {
            "description": "Already sorted list",
            "input": [1, 2, 3, 4, 5],
            "expected": [1, 2, 3, 4, 5],
        },
        {
            "description": "Reverse sorted list",
            "input": [5, 4, 3, 2, 1],
            "expected": [1, 2, 3, 4, 5],
        },
        {
            "description": "Single element",
            "input": [42],
            "expected": [42],
        },
        {
            "description": "Empty list",
            "input": [],
            "expected": [],
        },
        {
            "description": "List with duplicates",
            "input": [3, 1, 4, 1, 5, 9, 2, 6, 5, 3],
            "expected": [1, 1, 2, 3, 3, 4, 5, 5, 6, 9],
        },
        {
            "description": "Two elements",
            "input": [10, 7],
            "expected": [7, 10],
        },
    ]

    results = []
    all_passed = True

    for tc in test_cases:
        result = mergeSort(tc["input"][:])  # pass a copy to avoid mutation issues
        passed = result == tc["expected"]
        status = "PASS" if passed else "FAIL"
        if not passed:
            all_passed = False
        line = f"[{status}] {tc['description']}: mergeSort({tc['input']}) = {result}"
        results.append(line)
        print(line)

    print()
    if all_passed:
        summary = "All test cases passed!"
    else:
        summary = "Some test cases failed. Please check your implementation."
    print(summary)
    results.append("")
    results.append(summary)
    return results


if __name__ == "__main__":
    output_lines = run_tests()
    result_path = "/home/user/Desktop/result.txt"
    with open(result_path, "w") as f:
        f.write("\\n".join(output_lines) + "\\n")
    print(f"\\nOutput saved to {result_path}")
'''
    with open(STUB_PATH, 'w') as f:
        f.write(stub_code)
    print(f'Stub created: {STUB_PATH}')


def main():
    os.makedirs(DESKTOP, exist_ok=True)
    create_tutorial_docx()
    create_stub_py()

    # Open Merge_Sort_Tutorial.docx in LibreOffice Writer (task start state)
    launch_gui(f'libreoffice --writer "{TUTORIAL_PATH}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with Merge_Sort_Tutorial.docx (DISPLAY=:0)')


main()
