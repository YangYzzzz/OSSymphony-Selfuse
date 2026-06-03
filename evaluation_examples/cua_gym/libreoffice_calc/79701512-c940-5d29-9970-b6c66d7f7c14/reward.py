"""
Reward Script: Download PDF and check citation between research papers
Task ID: osworld_multi_apps_pdf_download_cite_003
Domain: multi_apps (os + libreoffice_calc + docx)

Scoring Rubric:
- Component 1: paper02.pdf exists as a valid PDF file (0.4 points)
- Component 2: cite_check.docx exists and contains 'Yes' answer (0.35 points)
- Component 3: cite_check.docx content references the correct papers (0.25 points)
Total: 1.0 points
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_pdf_download_cite_003'


def verify_task():
    """
    Verify task completion with progressive scoring.

    Task: Download the PDF of the 2nd paper (Playing Atari with Deep Reinforcement
    Learning, Mnih et al., 2013) and save as paper02.pdf in /home/user.
    Also record in cite_check.docx whether the 3rd paper (Nature 2015) cites the 2nd.
    Expected answer: Yes.

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    pdf_path = os.path.join(WORKDIR, 'paper02.pdf')
    docx_path = os.path.join(WORKDIR, 'cite_check.docx')

    # Component 1: paper02.pdf exists and is a valid PDF (0.4 points)
    # This must FAIL on initial_env (file doesn't exist) and PASS on golden_env
    try:
        if not os.path.isfile(pdf_path):
            print(f"FAIL: Component 1 — paper02.pdf not found at {pdf_path}")
        else:
            # Verify it is a genuine PDF by checking its magic bytes
            with open(pdf_path, 'rb') as f:
                header = f.read(8)
            if header.startswith(b'%PDF'):
                file_size = os.path.getsize(pdf_path)
                if file_size > 0:
                    print(f"PASS: Component 1 — paper02.pdf exists as valid PDF "
                          f"({file_size} bytes) (0.4 pts)")
                    total_score += 0.4
                else:
                    print("FAIL: Component 1 — paper02.pdf is empty (0 bytes)")
            else:
                print(f"FAIL: Component 1 — paper02.pdf exists but is not a valid PDF "
                      f"(header: {header[:8]})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: cite_check.docx exists and contains 'Yes' answer (0.35 points)
    # This must FAIL on initial_env and PASS on golden_env
    try:
        if not os.path.isfile(docx_path):
            print(f"FAIL: Component 2 — cite_check.docx not found at {docx_path}")
        else:
            from docx import Document
            doc = Document(docx_path)
            full_text = ' '.join(p.text.strip() for p in doc.paragraphs if p.text.strip())
            # Check for 'Yes' answer (case-insensitive) indicating citation confirmed
            # The answer must positively state the third paper cites the second
            if re.search(r'\byes\b', full_text, re.IGNORECASE):
                print(f"PASS: Component 2 — cite_check.docx contains 'Yes' answer (0.35 pts)")
                print(f"  Full text preview: {full_text[:200]}")
                total_score += 0.35
            else:
                print(f"FAIL: Component 2 — cite_check.docx does not contain 'Yes' answer")
                print(f"  Full text: {full_text[:300]}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: cite_check.docx references the correct papers in the citation check (0.25 points)
    # This verifies the content is about the correct papers (Atari/2013 and Nature/2015)
    # and not a generic 'Yes' answer for an unrelated question
    try:
        if not os.path.isfile(docx_path):
            print(f"FAIL: Component 3 — cite_check.docx not found")
        else:
            from docx import Document
            doc = Document(docx_path)
            full_text = ' '.join(p.text.strip() for p in doc.paragraphs if p.text.strip())
            full_text_lower = full_text.lower()

            # Check for references to the specific papers or citation context
            # The task involves: 2nd paper = Atari 2013, 3rd paper = Nature 2015
            has_citation_context = (
                # References to citation checking
                'cit' in full_text_lower or
                # References to the specific papers (Atari, DQN, reinforcement learning context)
                'atari' in full_text_lower or
                'reinforcement' in full_text_lower or
                'dqn' in full_text_lower or
                '2013' in full_text_lower or
                '2015' in full_text_lower or
                'mnih' in full_text_lower or
                # References to 'paper' or 'third' context
                ('paper' in full_text_lower and ('second' in full_text_lower or 'third' in full_text_lower))
            )

            if has_citation_context:
                print(f"PASS: Component 3 — cite_check.docx references relevant papers/context (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — cite_check.docx lacks context about the specific papers")
                print(f"  Full text: {full_text[:300]}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


verify_task()
