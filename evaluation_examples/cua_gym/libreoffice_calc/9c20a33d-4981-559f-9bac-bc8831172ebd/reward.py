"""
Reward Script: Save React Hooks documentation as a local Word file
Task ID: osworld_multi_apps_web_to_doc_007
Domain: multi_apps (web to docx)
Scoring:
  - Component 1: react_hooks.docx exists on the Desktop (gate only)
  - Component 2: Document contains 'Hooks' section heading and intro paragraphs (0.4 pts)
  - Component 3: Document lists specific hook names (useState, useEffect, useContext, etc.) (0.4 pts)
  - Component 4: Document contains sub-section headings for hook categories (0.2 pts)
Total: 1.0
"""

import os

# Use python-docx to verify the .docx file
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_web_to_doc_007'
FILE_PATH = '/home/user/Desktop/react_hooks.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition gate: file must exist
    if not os.path.exists(file_path):
        print(f"FAIL: File not found at {file_path}")
        print("REWARD: 0.0")
        return 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather all text for searching
    all_texts = [p.text.strip() for p in doc.paragraphs]
    all_text_combined = '\n'.join(all_texts)
    all_styles = [p.style.name for p in doc.paragraphs]

    # Component 1: Document contains 'Hooks' section heading and intro paragraphs (0.4 pts)
    # The task requires saving the page from intro through the list of hooks.
    # The document should contain a 'Hooks' heading and introductory React content.
    try:
        hooks_heading_paras = [
            p for p in doc.paragraphs
            if 'hooks' in p.text.lower().strip()
            and p.style.name in ('Heading 1', 'Heading 2', 'Heading 3', 'Normal')
        ]
        intro_paras = [
            p for p in doc.paragraphs
            if 'react' in p.text.lower().strip() and len(p.text.strip()) > 20
        ]

        if hooks_heading_paras and intro_paras:
            print(f"PASS: Component 1 — 'Hooks' section heading and React intro text found (0.4 pts)")
            total_score += 0.4
        elif hooks_heading_paras:
            print(f"PARTIAL: Component 1 — 'Hooks' heading found but no React intro text (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 1 — No 'Hooks' section heading found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Document lists specific hook names (useState, useEffect, useContext, etc.) (0.4 pts)
    # The React Hooks reference page lists built-in hooks. We verify key hooks are present.
    try:
        key_hooks = ['useState', 'useEffect', 'useContext', 'useReducer', 'useRef', 'useMemo', 'useCallback']
        found_hooks = []
        for hook in key_hooks:
            if hook in all_text_combined:
                found_hooks.append(hook)

        hook_ratio = len(found_hooks) / len(key_hooks)

        if hook_ratio >= 0.7:
            print(f"PASS: Component 2 — Found {len(found_hooks)}/{len(key_hooks)} key hooks: {found_hooks} (0.4 pts)")
            total_score += 0.4
        elif hook_ratio >= 0.4:
            print(f"PARTIAL: Component 2 — Found only {len(found_hooks)}/{len(key_hooks)} key hooks: {found_hooks} (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 2 — Found only {len(found_hooks)}/{len(key_hooks)} key hooks: {found_hooks}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Document contains sub-section headings for hook categories (0.2 pts)
    # The page has structured sections: State Hooks, Context Hooks, Ref Hooks, Effect Hooks, etc.
    try:
        hook_categories = ['State Hooks', 'Context Hooks', 'Ref Hooks', 'Effect Hooks', 'Performance Hooks']
        found_categories = []
        for cat in hook_categories:
            if cat in all_text_combined:
                found_categories.append(cat)

        category_ratio = len(found_categories) / len(hook_categories)

        if category_ratio >= 0.6:
            print(f"PASS: Component 3 — Found {len(found_categories)}/{len(hook_categories)} hook category sections: {found_categories} (0.2 pts)")
            total_score += 0.2
        elif category_ratio >= 0.3:
            print(f"PARTIAL: Component 3 — Found {len(found_categories)}/{len(hook_categories)} hook category sections: {found_categories} (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 3 — Found only {len(found_categories)}/{len(hook_categories)} hook category sections: {found_categories}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Run verification
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
