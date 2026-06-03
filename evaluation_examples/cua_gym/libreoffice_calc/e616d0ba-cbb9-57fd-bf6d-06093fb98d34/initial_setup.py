"""
Initial Setup: Open LibreOffice Writer with a networking tutorial document.
Task ID: osworld_multi_apps_terminal_screenshot_003
Domain: multi_apps (LibreOffice Writer + Terminal + Screenshot)

Initial state:
  - LibreOffice Writer open with a networking tutorial document
  - No terminal window open
  - No network_info.png on the Desktop
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_terminal_screenshot_003'
DESKTOP = f'{WORKDIR}/Desktop'
DOC_PATH = f'{WORKDIR}/{TASK_ID}.odt'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_networking_tutorial():
    """Create a networking tutorial document using LibreOffice UNO macro via python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        subprocess.run(['pip3', 'install', 'python-docx'], check=True)
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading('Linux Networking Tutorial', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_heading('Introduction', level=2)
    intro = doc.add_paragraph(
        'Understanding network configuration is a fundamental skill for any Linux administrator. '
        'This tutorial covers the essential tools and commands for viewing and managing network '
        'interfaces on a Linux system.'
    )

    # Section 1: Network Interfaces
    doc.add_heading('1. Network Interfaces', level=2)
    doc.add_paragraph(
        'A network interface is the point of interconnection between a computer and a network. '
        'On Linux systems, network interfaces can be physical (Ethernet cards, Wi-Fi adapters) '
        'or virtual (loopback, tunnels, bridges).'
    )

    doc.add_heading('Common Interface Types', level=3)
    doc.add_paragraph('eth0 / enp0s3  — Wired Ethernet interface', style='List Bullet')
    doc.add_paragraph('wlan0 / wlp2s0 — Wireless LAN interface', style='List Bullet')
    doc.add_paragraph('lo             — Loopback interface (127.0.0.1)', style='List Bullet')
    doc.add_paragraph('docker0        — Docker virtual bridge interface', style='List Bullet')

    # Section 2: The ifconfig Command
    doc.add_heading('2. The ifconfig Command', level=2)
    doc.add_paragraph(
        'The ifconfig (interface configuration) command is a traditional Unix/Linux tool used to '
        'configure and display network interface parameters. Although newer systems prefer the '
        'ip command, ifconfig remains widely used and is part of the net-tools package.'
    )

    doc.add_heading('Basic Syntax', level=3)
    code_para = doc.add_paragraph('$ ifconfig [interface] [options]')
    code_para.runs[0].font.name = 'Courier New'
    code_para.runs[0].font.size = Pt(10)

    doc.add_heading('Common Usage', level=3)
    doc.add_paragraph('Display all active interfaces:', style='List Number')
    p = doc.add_paragraph('$ ifconfig')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph('Display all interfaces (including inactive):', style='List Number')
    p = doc.add_paragraph('$ ifconfig -a')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph('Display a specific interface:', style='List Number')
    p = doc.add_paragraph('$ ifconfig eth0')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)

    # Section 3: Reading ifconfig Output
    doc.add_heading('3. Reading ifconfig Output', level=2)
    doc.add_paragraph(
        'When you run ifconfig, the output provides detailed information about each network interface. '
        'Understanding this output is crucial for network troubleshooting and configuration.'
    )

    doc.add_heading('Key Fields in ifconfig Output', level=3)
    fields = [
        ('inet', 'The IPv4 address assigned to the interface'),
        ('inet6', 'The IPv6 address assigned to the interface'),
        ('netmask', 'The subnet mask defining the network range'),
        ('broadcast', 'The broadcast address for the network'),
        ('ether', 'The MAC (hardware) address of the interface'),
        ('RX packets', 'Number of packets received'),
        ('TX packets', 'Number of packets transmitted'),
        ('MTU', 'Maximum Transmission Unit — largest packet size allowed'),
    ]
    for field, desc in fields:
        p = doc.add_paragraph(style='List Bullet')
        run_bold = p.add_run(f'{field}: ')
        run_bold.bold = True
        p.add_run(desc)

    # Section 4: Practical Exercise
    doc.add_heading('4. Practical Exercise', level=2)
    doc.add_paragraph(
        'Now it is your turn to explore the network configuration of this system. '
        'Follow the steps below to capture the output of the ifconfig command:'
    )

    steps = [
        'Open a terminal emulator (e.g., GNOME Terminal or xterm).',
        "Type the command 'ifconfig' and press Enter.",
        'Review the output, noting each active network interface.',
        'Take a screenshot of the terminal window showing the ifconfig output.',
        "Save the screenshot as 'network_info.png' to the Desktop.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')

    # Section 5: Modern Alternative
    doc.add_heading('5. Modern Alternative: ip Command', level=2)
    doc.add_paragraph(
        'On modern Linux distributions, the ip command from the iproute2 package has largely '
        'replaced ifconfig. However, ifconfig is still available and commonly used in scripts '
        'and legacy systems.'
    )
    p = doc.add_paragraph('# Equivalent commands:')
    p.runs[0].font.name = 'Courier New'
    p = doc.add_paragraph('$ ifconfig          →  $ ip addr show')
    p.runs[0].font.name = 'Courier New'
    p = doc.add_paragraph('$ ifconfig eth0     →  $ ip addr show eth0')
    p.runs[0].font.name = 'Courier New'

    # References
    doc.add_heading('References', level=2)
    doc.add_paragraph('man ifconfig — Manual page for ifconfig', style='List Bullet')
    doc.add_paragraph('net-tools documentation: https://net-tools.sourceforge.io/', style='List Bullet')
    doc.add_paragraph('Linux Network Administrators Guide', style='List Bullet')

    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    # Remove any existing network_info.png to ensure clean initial state
    png_path = f'{DESKTOP}/network_info.png'
    if os.path.exists(png_path):
        os.remove(png_path)
        print(f'Removed existing {png_path} to ensure clean initial state')

    # Save as .odt via python-docx (saves as .docx, rename to .odt for LibreOffice)
    # Actually save as .docx and open with LibreOffice Writer
    doc_path_docx = f'{WORKDIR}/{TASK_ID}.docx'
    doc.save(doc_path_docx)
    print(f'Networking tutorial document created: {doc_path_docx}')

    return doc_path_docx


def create_initial():
    doc_path = create_networking_tutorial()

    # Launch LibreOffice Writer with the tutorial document
    launch_gui(f'libreoffice --writer "{doc_path}"', delay_sec=3.0)
    print(f'GUI_READY: LibreOffice Writer launched with networking tutorial at DISPLAY=:0')
    print(f'Initial state: Writer open, no terminal, no network_info.png on Desktop')


create_initial()
