"""
Initial Setup: Merge 6 different-format documents into a single PDF
Task ID: pdf_gf3_018
Domain: pdf (libreoffice_calc tagged but actually PDF merge task)

Creates all source files that the agent needs to merge:
  - /home/user/docs/cover.docx
  - /home/user/docs/section_a.pdf
  - /home/user/docs/section_b.pdf
  - /home/user/docs/section_c.pdf
  - /home/user/images/signature.png
Opens a file manager so the agent can see the workspace.
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'pdf_gf3_018'
DOCS_DIR = f'{WORKDIR}/docs'
IMAGES_DIR = f'{WORKDIR}/images'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_cover_docx():
    """Create a realistic cover letter DOCX."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading('Government Submission Cover Letter', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    # Sender info
    p = doc.add_paragraph()
    p.add_run('Meridian Consulting Group\n').bold = True
    p.add_run('1450 Pennsylvania Avenue NW, Suite 320\n')
    p.add_run('Washington, DC 20004\n')
    p.add_run('Tel: (202) 555-0187 | Fax: (202) 555-0188\n')
    p.add_run('Email: submissions@meridiancg.com')

    doc.add_paragraph('')

    # Date and recipient
    doc.add_paragraph('Date: March 28, 2025')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.add_run('Office of Procurement Services\n').bold = True
    p.add_run('U.S. Department of Energy\n')
    p.add_run('1000 Independence Avenue SW\n')
    p.add_run('Washington, DC 20585')

    doc.add_paragraph('')

    # Subject
    p = doc.add_paragraph()
    p.add_run('RE: ').bold = True
    p.add_run('Response to Solicitation No. DOE-RFP-2025-0347 - ')
    p.add_run('Advanced Energy Grid Modernization Assessment')

    doc.add_paragraph('')

    # Body
    doc.add_paragraph(
        'Dear Procurement Officer,'
    )
    doc.add_paragraph(
        'Meridian Consulting Group is pleased to submit the enclosed proposal in response '
        'to the above-referenced solicitation for Advanced Energy Grid Modernization Assessment '
        'services. Our firm has over 18 years of experience in energy infrastructure consulting, '
        'having successfully completed 47 federal engagements across the Department of Energy, '
        'Department of Defense, and the Environmental Protection Agency.'
    )
    doc.add_paragraph(
        'This submission package includes the following sections: (A) Technical Approach, '
        '(B) Past Performance References, (C) Cost Proposal, and a title page summarizing '
        'the key personnel assigned to this project. A signed authorization page is included '
        'as the final page of this document.'
    )
    doc.add_paragraph(
        'We confirm that all information provided herein is accurate and complete to the best '
        'of our knowledge. Meridian Consulting Group is fully committed to meeting all requirements '
        'set forth in the solicitation and stands ready to begin work upon contract award.'
    )
    doc.add_paragraph('')

    # Closing
    doc.add_paragraph('Respectfully submitted,')
    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.add_run('Dr. Evelyn R. Hartwell\n').bold = True
    p.add_run('Managing Director\n')
    p.add_run('Meridian Consulting Group')

    doc.save(f'{DOCS_DIR}/cover.docx')
    print(f'Created: {DOCS_DIR}/cover.docx')


def create_section_a_pdf():
    """Create Section A: Technical Approach (2 pages)."""
    import pymupdf

    doc = pymupdf.open()

    # Page 1
    page = doc.new_page(width=595, height=842)
    y = 60

    page.insert_text(pymupdf.Point(72, y), "SECTION A: TECHNICAL APPROACH",
                     fontsize=18, fontname="hebo", color=(0, 0, 0.5))
    y += 40

    page.insert_text(pymupdf.Point(72, y), "1. Overview",
                     fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    text_block = (
        "Meridian Consulting Group proposes a comprehensive three-phase approach to the "
        "Advanced Energy Grid Modernization Assessment. Our methodology integrates advanced "
        "analytics, field assessments, and stakeholder engagement to deliver actionable "
        "recommendations for grid infrastructure improvements across the targeted regions."
    )
    rect = pymupdf.Rect(72, y, 523, y + 80)
    page.insert_textbox(rect, text_block, fontsize=11, fontname="helv",
                        align=pymupdf.TEXT_ALIGN_JUSTIFY)
    y += 95

    page.insert_text(pymupdf.Point(72, y), "2. Phase I - Assessment and Data Collection",
                     fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    phase1_text = (
        "During Phase I (Months 1-4), our team will conduct a thorough assessment of existing "
        "grid infrastructure across the 12 designated utility districts. This includes: on-site "
        "inspections of 340+ substations, analysis of 15 years of historical performance data, "
        "interviews with 85 regional utility operators, and deployment of advanced sensor arrays "
        "at 48 critical junction points. Our proprietary GridSense analytics platform will process "
        "over 2.3 million data points to establish baseline performance metrics."
    )
    rect = pymupdf.Rect(72, y, 523, y + 110)
    page.insert_textbox(rect, phase1_text, fontsize=11, fontname="helv",
                        align=pymupdf.TEXT_ALIGN_JUSTIFY)
    y += 125

    page.insert_text(pymupdf.Point(72, y), "3. Phase II - Analysis and Modeling",
                     fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    phase2_text = (
        "Phase II (Months 5-8) focuses on comprehensive data analysis and predictive modeling. "
        "Using our GridFuture simulation engine, we will model 24 different modernization scenarios "
        "incorporating renewable energy integration, demand response optimization, and resilience "
        "enhancements. Each scenario will be evaluated against 18 key performance indicators "
        "including reliability indices (SAIDI, SAIFI, CAIDI), cost-effectiveness ratios, "
        "carbon emission reductions, and cybersecurity readiness scores."
    )
    rect = pymupdf.Rect(72, y, 523, y + 110)
    page.insert_textbox(rect, phase2_text, fontsize=11, fontname="helv",
                        align=pymupdf.TEXT_ALIGN_JUSTIFY)
    y += 125

    page.insert_text(pymupdf.Point(72, y), "4. Phase III - Recommendations and Reporting",
                     fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    phase3_text = (
        "In Phase III (Months 9-12), Meridian will synthesize findings into a comprehensive "
        "modernization roadmap. Deliverables include: a 200+ page technical report, an executive "
        "summary briefing package, a prioritized investment schedule spanning 10 years, and "
        "interactive data dashboards accessible via secure web portal."
    )
    rect = pymupdf.Rect(72, y, 523, y + 80)
    page.insert_textbox(rect, phase3_text, fontsize=11, fontname="helv",
                        align=pymupdf.TEXT_ALIGN_JUSTIFY)

    # Footer
    page.insert_text(pymupdf.Point(280, 810), "A-1", fontsize=9, fontname="helv", color=(0.5, 0.5, 0.5))

    # Page 2
    page2 = doc.new_page(width=595, height=842)
    y = 60

    page2.insert_text(pymupdf.Point(72, y), "5. Key Personnel",
                      fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    personnel_text = (
        "Project Director: Dr. Evelyn R. Hartwell, PE - 22 years experience in energy "
        "infrastructure assessment, former Senior Advisor at DOE Office of Electricity.\n\n"
        "Lead Engineer: Marcus J. Okafor, PhD - Specialist in grid resilience modeling "
        "with 15 years experience and 34 peer-reviewed publications.\n\n"
        "Data Analytics Lead: Dr. Priya Ramanathan - Expert in machine learning applications "
        "for utility operations, developed the GridSense analytics platform.\n\n"
        "Field Operations Manager: Captain (Ret.) Thomas W. Brennan - 18 years managing "
        "large-scale infrastructure inspection programs across federal installations."
    )
    rect = pymupdf.Rect(72, y, 523, y + 180)
    page2.insert_textbox(rect, personnel_text, fontsize=11, fontname="helv",
                         align=pymupdf.TEXT_ALIGN_LEFT)
    y += 200

    page2.insert_text(pymupdf.Point(72, y), "6. Quality Assurance",
                      fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    qa_text = (
        "Meridian maintains ISO 9001:2015 certification and implements a rigorous quality "
        "management system for all deliverables. Our QA process includes: independent peer "
        "review of all technical analyses, automated data validation checks at each processing "
        "stage, monthly progress reviews with DOE program management, and final deliverable "
        "review by our Technical Advisory Board comprising three former DOE Deputy Secretaries."
    )
    rect = pymupdf.Rect(72, y, 523, y + 100)
    page2.insert_textbox(rect, qa_text, fontsize=11, fontname="helv",
                         align=pymupdf.TEXT_ALIGN_JUSTIFY)
    y += 115

    page2.insert_text(pymupdf.Point(72, y), "7. Schedule and Milestones",
                      fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    schedule_text = (
        "Month 1: Project kickoff and team mobilization\n"
        "Month 2-4: Field assessments and data collection\n"
        "Month 5-6: Data analysis and scenario development\n"
        "Month 7-8: Modeling and simulation runs\n"
        "Month 9-10: Draft report preparation\n"
        "Month 11: Stakeholder review period\n"
        "Month 12: Final report delivery and closeout briefing"
    )
    rect = pymupdf.Rect(72, y, 523, y + 120)
    page2.insert_textbox(rect, schedule_text, fontsize=11, fontname="helv",
                         align=pymupdf.TEXT_ALIGN_LEFT)

    # Footer
    page2.insert_text(pymupdf.Point(280, 810), "A-2", fontsize=9, fontname="helv", color=(0.5, 0.5, 0.5))

    doc.save(f'{DOCS_DIR}/section_a.pdf')
    doc.close()
    print(f'Created: {DOCS_DIR}/section_a.pdf')


def create_section_b_pdf():
    """Create Section B: Past Performance (1 page)."""
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page(width=595, height=842)
    y = 60

    page.insert_text(pymupdf.Point(72, y), "SECTION B: PAST PERFORMANCE REFERENCES",
                     fontsize=18, fontname="hebo", color=(0, 0, 0.5))
    y += 40

    page.insert_text(pymupdf.Point(72, y), "Reference 1: DOE Grid Resilience Study (2022-2024)",
                     fontsize=12, fontname="hebo", color=(0, 0, 0))
    y += 20
    ref1 = (
        "Contract No. DE-AC05-22OR35470 | Value: $4.2M | Rating: Exceptional\n"
        "Conducted comprehensive resilience assessment of the Southeast Regional Grid "
        "covering 8 states and 23 utility providers. Delivered final report 3 weeks ahead "
        "of schedule with zero corrective action requests."
    )
    rect = pymupdf.Rect(72, y, 523, y + 70)
    page.insert_textbox(rect, ref1, fontsize=10, fontname="helv")
    y += 85

    page.insert_text(pymupdf.Point(72, y), "Reference 2: EPA Facility Energy Audit Program (2021-2023)",
                     fontsize=12, fontname="hebo", color=(0, 0, 0))
    y += 20
    ref2 = (
        "Contract No. EP-W-21-004 | Value: $2.8M | Rating: Very Good\n"
        "Performed energy audits of 156 EPA-owned facilities nationwide. Identified "
        "$18.5M in potential annual energy savings. Recommendations led to 23% reduction "
        "in aggregate energy consumption within 18 months of implementation."
    )
    rect = pymupdf.Rect(72, y, 523, y + 70)
    page.insert_textbox(rect, ref2, fontsize=10, fontname="helv")
    y += 85

    page.insert_text(pymupdf.Point(72, y), "Reference 3: DoD Installation Modernization Review (2020-2022)",
                     fontsize=12, fontname="hebo", color=(0, 0, 0))
    y += 20
    ref3 = (
        "Contract No. W912DY-20-C-0089 | Value: $5.7M | Rating: Exceptional\n"
        "Led multi-disciplinary team assessing power infrastructure at 34 military "
        "installations across CONUS. Developed prioritized $890M investment roadmap "
        "adopted by the Office of the Deputy Assistant Secretary of Defense for Installations."
    )
    rect = pymupdf.Rect(72, y, 523, y + 70)
    page.insert_textbox(rect, ref3, fontsize=10, fontname="helv")
    y += 85

    page.insert_text(pymupdf.Point(72, y), "Reference 4: State of California Grid Integration Study (2023-2024)",
                     fontsize=12, fontname="hebo", color=(0, 0, 0))
    y += 20
    ref4 = (
        "Contract No. CA-CPUC-2023-0156 | Value: $3.1M | Rating: Outstanding\n"
        "Analyzed integration pathways for 15 GW of renewable energy into the California "
        "grid. Developed dynamic simulation models predicting grid behavior under 200+ "
        "scenarios. Report cited in CPUC Decision D.24-03-028."
    )
    rect = pymupdf.Rect(72, y, 523, y + 70)
    page.insert_textbox(rect, ref4, fontsize=10, fontname="helv")

    # Footer
    page.insert_text(pymupdf.Point(280, 810), "B-1", fontsize=9, fontname="helv", color=(0.5, 0.5, 0.5))

    doc.save(f'{DOCS_DIR}/section_b.pdf')
    doc.close()
    print(f'Created: {DOCS_DIR}/section_b.pdf')


def create_section_c_pdf():
    """Create Section C: Cost Proposal (2 pages)."""
    import pymupdf

    doc = pymupdf.open()

    # Page 1
    page = doc.new_page(width=595, height=842)
    y = 60

    page.insert_text(pymupdf.Point(72, y), "SECTION C: COST PROPOSAL",
                     fontsize=18, fontname="hebo", color=(0, 0, 0.5))
    y += 40

    page.insert_text(pymupdf.Point(72, y), "1. Cost Summary",
                     fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    summary = (
        "The total proposed cost for the Advanced Energy Grid Modernization Assessment "
        "is $6,847,500.00 over a 12-month period of performance. This cost includes all "
        "labor, travel, equipment, materials, and indirect costs necessary to complete "
        "the scope of work described in Section A."
    )
    rect = pymupdf.Rect(72, y, 523, y + 70)
    page.insert_textbox(rect, summary, fontsize=11, fontname="helv",
                        align=pymupdf.TEXT_ALIGN_JUSTIFY)
    y += 85

    page.insert_text(pymupdf.Point(72, y), "2. Labor Cost Breakdown",
                     fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    # Table header
    shape = page.new_shape()
    # Header row background
    header_rect = pymupdf.Rect(72, y, 523, y + 20)
    shape.draw_rect(header_rect)
    shape.finish(fill=(0, 0, 0.5), color=(0, 0, 0.5))
    shape.commit()

    page.insert_text(pymupdf.Point(76, y + 14), "Position", fontsize=10, fontname="hebo", color=(1, 1, 1))
    page.insert_text(pymupdf.Point(220, y + 14), "Hours", fontsize=10, fontname="hebo", color=(1, 1, 1))
    page.insert_text(pymupdf.Point(290, y + 14), "Rate ($/hr)", fontsize=10, fontname="hebo", color=(1, 1, 1))
    page.insert_text(pymupdf.Point(400, y + 14), "Total", fontsize=10, fontname="hebo", color=(1, 1, 1))
    y += 22

    rows = [
        ("Project Director", "1,200", "$285.00", "$342,000"),
        ("Lead Engineer", "2,400", "$245.00", "$588,000"),
        ("Data Analytics Lead", "2,100", "$230.00", "$483,000"),
        ("Field Operations Mgr", "1,800", "$195.00", "$351,000"),
        ("Senior Engineers (4)", "7,200", "$185.00", "$1,332,000"),
        ("Junior Engineers (6)", "9,600", "$125.00", "$1,200,000"),
        ("Data Analysts (3)", "4,800", "$140.00", "$672,000"),
        ("Administrative Support", "2,400", "$85.00", "$204,000"),
    ]

    for i, (pos, hrs, rate, total) in enumerate(rows):
        bg = (0.95, 0.95, 1.0) if i % 2 == 0 else (1, 1, 1)
        shape = page.new_shape()
        row_rect = pymupdf.Rect(72, y, 523, y + 18)
        shape.draw_rect(row_rect)
        shape.finish(fill=bg, color=(0.7, 0.7, 0.7))
        shape.commit()

        page.insert_text(pymupdf.Point(76, y + 13), pos, fontsize=9, fontname="helv")
        page.insert_text(pymupdf.Point(225, y + 13), hrs, fontsize=9, fontname="helv")
        page.insert_text(pymupdf.Point(295, y + 13), rate, fontsize=9, fontname="helv")
        page.insert_text(pymupdf.Point(405, y + 13), total, fontsize=9, fontname="helv")
        y += 18

    # Total row
    y += 5
    shape = page.new_shape()
    shape.draw_line(pymupdf.Point(72, y), pymupdf.Point(523, y))
    shape.finish(color=(0, 0, 0), width=1.5)
    shape.commit()
    y += 15
    page.insert_text(pymupdf.Point(76, y), "Subtotal - Direct Labor", fontsize=10, fontname="hebo")
    page.insert_text(pymupdf.Point(400, y), "$5,172,000", fontsize=10, fontname="hebo")

    y += 35
    page.insert_text(pymupdf.Point(72, y), "3. Other Direct Costs",
                     fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    odc_items = [
        ("Travel (domestic, 48 trips)", "$384,000"),
        ("Equipment and sensors", "$425,500"),
        ("Software licenses", "$186,000"),
        ("Subcontractor support", "$312,000"),
        ("Printing and reproduction", "$18,000"),
    ]
    for item, cost in odc_items:
        page.insert_text(pymupdf.Point(90, y + 13), item, fontsize=10, fontname="helv")
        page.insert_text(pymupdf.Point(405, y + 13), cost, fontsize=10, fontname="helv")
        y += 18

    y += 5
    shape = page.new_shape()
    shape.draw_line(pymupdf.Point(72, y), pymupdf.Point(523, y))
    shape.finish(color=(0, 0, 0), width=1)
    shape.commit()
    y += 15
    page.insert_text(pymupdf.Point(76, y), "Subtotal - ODCs", fontsize=10, fontname="hebo")
    page.insert_text(pymupdf.Point(400, y), "$1,325,500", fontsize=10, fontname="hebo")

    # Footer
    page.insert_text(pymupdf.Point(280, 810), "C-1", fontsize=9, fontname="helv", color=(0.5, 0.5, 0.5))

    # Page 2
    page2 = doc.new_page(width=595, height=842)
    y = 60

    page2.insert_text(pymupdf.Point(72, y), "4. Indirect Costs and Fee",
                      fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    indirect_items = [
        ("Overhead (12% of direct labor)", "$620,640"),
        ("General & Administrative (8%)", "$413,760"),
        ("Fee (5% of total costs)", "$315,600"),
    ]
    for item, cost in indirect_items:
        page2.insert_text(pymupdf.Point(90, y + 13), item, fontsize=10, fontname="helv")
        page2.insert_text(pymupdf.Point(405, y + 13), cost, fontsize=10, fontname="helv")
        y += 20

    y += 10
    shape = page2.new_shape()
    shape.draw_line(pymupdf.Point(72, y), pymupdf.Point(523, y))
    shape.finish(color=(0, 0, 0), width=2)
    shape.commit()
    y += 20

    page2.insert_text(pymupdf.Point(76, y), "TOTAL PROPOSED COST", fontsize=14, fontname="hebo",
                      color=(0, 0, 0.5))
    page2.insert_text(pymupdf.Point(380, y), "$6,847,500.00", fontsize=14, fontname="hebo",
                      color=(0, 0, 0.5))

    y += 50
    page2.insert_text(pymupdf.Point(72, y), "5. Payment Schedule",
                      fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    payment_text = (
        "Meridian proposes monthly invoicing based on actual costs incurred plus applicable "
        "indirect rates and fee. Invoices will be submitted by the 15th of each month for "
        "the preceding month's costs. Payment terms: Net 30 from receipt of properly prepared "
        "invoice."
    )
    rect = pymupdf.Rect(72, y, 523, y + 70)
    page2.insert_textbox(rect, payment_text, fontsize=11, fontname="helv",
                         align=pymupdf.TEXT_ALIGN_JUSTIFY)
    y += 85

    page2.insert_text(pymupdf.Point(72, y), "6. Period of Performance",
                      fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25

    pop_text = (
        "Base Period: 12 months from date of contract award\n"
        "Option Period 1: 6 months (follow-on analysis and support)\n"
        "Option Period 2: 6 months (implementation oversight)\n\n"
        "All option periods are priced separately and will be exercised at the "
        "Government's discretion."
    )
    rect = pymupdf.Rect(72, y, 523, y + 100)
    page2.insert_textbox(rect, pop_text, fontsize=11, fontname="helv",
                         align=pymupdf.TEXT_ALIGN_LEFT)

    y += 120
    page2.insert_text(pymupdf.Point(72, y), "7. Validity",
                      fontsize=14, fontname="hebo", color=(0, 0, 0))
    y += 25
    page2.insert_text(pymupdf.Point(72, y + 5),
                      "This cost proposal is valid for 120 days from the date of submission.",
                      fontsize=11, fontname="helv")

    # Footer
    page2.insert_text(pymupdf.Point(280, 810), "C-2", fontsize=9, fontname="helv", color=(0.5, 0.5, 0.5))

    doc.save(f'{DOCS_DIR}/section_c.pdf')
    doc.close()
    print(f'Created: {DOCS_DIR}/section_c.pdf')


def create_signature_png():
    """Create a realistic signature image."""
    from PIL import Image, ImageDraw, ImageFont

    # Create white background at reasonable resolution
    img = Image.new('RGB', (800, 400), 'white')
    draw = ImageDraw.Draw(img)

    # Draw a signature-like curved line (simulated handwriting)
    # Using series of connected arcs/curves
    points = [
        (100, 250), (120, 200), (150, 180), (180, 220), (200, 260),
        (220, 200), (250, 160), (280, 180), (300, 220), (320, 240),
        (350, 200), (380, 170), (400, 190), (420, 230), (440, 250),
        (460, 200), (490, 180), (520, 210), (550, 230), (570, 220),
        (600, 200), (630, 210), (660, 230), (680, 220), (700, 210),
    ]
    draw.line(points, fill='navy', width=3)

    # Add a second stroke for realism
    points2 = [
        (150, 270), (200, 265), (250, 275), (300, 268), (350, 272),
        (400, 268), (450, 275), (500, 270),
    ]
    draw.line(points2, fill='navy', width=2)

    # Add printed name below
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 24)
    except (OSError, IOError):
        font = ImageFont.load_default()

    draw.text((200, 310), "Dr. Evelyn R. Hartwell", fill='black', font=font)
    draw.text((200, 340), "Managing Director", fill='gray', font=font)

    # Add a line above signature
    draw.line([(100, 290), (700, 290)], fill='black', width=1)

    img.save(f'{IMAGES_DIR}/signature.png')
    print(f'Created: {IMAGES_DIR}/signature.png')


def create_initial():
    """Create all source files for the merge task."""
    os.makedirs(DOCS_DIR, exist_ok=True)
    os.makedirs(IMAGES_DIR, exist_ok=True)

    create_cover_docx()
    create_section_a_pdf()
    create_section_b_pdf()
    create_section_c_pdf()
    create_signature_png()

    print(f'\nAll source files created in {DOCS_DIR}/ and {IMAGES_DIR}/')
    print('Source files:')
    print(f'  {DOCS_DIR}/cover.docx')
    print(f'  {DOCS_DIR}/section_a.pdf')
    print(f'  {DOCS_DIR}/section_b.pdf')
    print(f'  {DOCS_DIR}/section_c.pdf')
    print(f'  {IMAGES_DIR}/signature.png')

    # GUI-ready: open file manager showing docs directory
    launch_gui(f'nautilus "{DOCS_DIR}"', delay_sec=2.0)
    print('GUI_READY: launched Nautilus file manager with DISPLAY=:0')


create_initial()
