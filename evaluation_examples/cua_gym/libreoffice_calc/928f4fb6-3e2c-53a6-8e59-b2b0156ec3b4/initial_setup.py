"""
Initial Setup: macro_requirements.docx + report_automation.xlsx with raw transaction data
Task ID: osworld_multi_apps_docx_to_calc_013
Domain: libreoffice_calc (multi-app: also creates a docx)
"""

import os
import shlex
import subprocess
import time
import datetime

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_docx_to_calc_013'
XLSX_OUTPUT = f'{WORKDIR}/report_automation.xlsx'
DOCX_OUTPUT = f'{WORKDIR}/Desktop/macro_requirements.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_macro_requirements_docx():
    """Create macro_requirements.docx on the Desktop describing the 3 required macros."""
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading('Macro Requirements for report_automation.xlsx', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction
    doc.add_paragraph(
        'This document specifies three LibreOffice Basic macros to be implemented '
        'in the report_automation.xlsx workbook. Each macro should be accessible '
        'via the Tools > Macros menu in LibreOffice Calc.'
    )

    doc.add_paragraph('')  # blank line

    # Macro 1
    h1 = doc.add_heading('Macro 1: SortByDate', level=1)
    doc.add_paragraph(
        'Purpose: Automatically sort the transaction data on Sheet1 by the Date '
        'column (column A) in ascending order.'
    )
    doc.add_paragraph('Requirements:')
    doc.add_paragraph('- Sort all data rows (rows 2 onward) by column A (Date) in ascending order.', style='List Bullet')
    doc.add_paragraph('- Keep header row (row 1) in place.', style='List Bullet')
    doc.add_paragraph('- Macro name: SortByDate', style='List Bullet')
    doc.add_paragraph('- After sorting, save the workbook.', style='List Bullet')

    doc.add_paragraph('')

    # Macro 2
    doc.add_heading('Macro 2: GenerateMonthlySummary', level=1)
    doc.add_paragraph(
        'Purpose: Generate a monthly summary table on Sheet2 based on the '
        'transaction data in Sheet1.'
    )
    doc.add_paragraph('Requirements:')
    doc.add_paragraph('- Read all transaction data from Sheet1 (columns: Date, Description, Amount, Category).', style='List Bullet')
    doc.add_paragraph('- Group transactions by month (format: YYYY-MM) and calculate the total Amount for each month.', style='List Bullet')
    doc.add_paragraph('- Write summary to Sheet2 with headers: Month, Total Amount, Transaction Count.', style='List Bullet')
    doc.add_paragraph('- Sort summary rows by Month in ascending order.', style='List Bullet')
    doc.add_paragraph('- Macro name: GenerateMonthlySummary', style='List Bullet')
    doc.add_paragraph('- After generating the summary, save the workbook.', style='List Bullet')

    doc.add_paragraph('')

    # Macro 3
    doc.add_heading('Macro 3: ExportSummaryAsPDF', level=1)
    doc.add_paragraph(
        'Purpose: Export the monthly summary table on Sheet2 as a PDF file '
        'named "monthly_summary.pdf" to the Desktop.'
    )
    doc.add_paragraph('Requirements:')
    doc.add_paragraph('- Activate Sheet2.', style='List Bullet')
    doc.add_paragraph('- Export Sheet2 content as a PDF file.', style='List Bullet')
    doc.add_paragraph('- Save the PDF to: /home/user/Desktop/monthly_summary.pdf', style='List Bullet')
    doc.add_paragraph('- Macro name: ExportSummaryAsPDF', style='List Bullet')
    doc.add_paragraph('- After export, save the workbook.', style='List Bullet')

    doc.add_paragraph('')

    # Implementation notes
    doc.add_heading('Implementation Notes', level=1)
    doc.add_paragraph(
        'All three macros should be written as LibreOffice Basic subroutines '
        'stored in the workbook\'s Basic module (not in the application library). '
        'They should be accessible from Tools > Macros > Organize Basic Macros.'
    )

    doc.save(DOCX_OUTPUT)
    print(f'Macro requirements document created: {DOCX_OUTPUT}')


def create_report_automation_xlsx():
    """Create report_automation.xlsx with raw transaction data on Sheet1, empty Sheet2."""
    wb = openpyxl.Workbook()

    # --- Sheet 1: Transactions (raw data, NOT sorted by date) ---
    ws1 = wb.active
    ws1.title = 'Sheet1'

    # Header row
    headers = ['Date', 'Description', 'Amount', 'Category']
    header_font = Font(name='Calibri', bold=True, size=11)
    header_fill = PatternFill(start_color='FF4472C4', end_color='FF4472C4', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center')
    thin = Side(style='thin', color='000000')
    header_border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, h in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=col, value=h)
        cell.font = Font(name='Calibri', bold=True, size=11, color='FFFFFFFF')
        cell.fill = PatternFill(start_color='FF4472C4', end_color='FF4472C4', fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # Transaction data — deliberately NOT sorted by date (mixed order)
    # Using realistic business transaction data
    transactions = [
        ('2025-03-15', 'Office Supplies - Staples', 234.50, 'Office'),
        ('2025-01-08', 'Cloud Hosting - AWS', 1850.00, 'IT'),
        ('2025-04-22', 'Team Lunch - Bistro Milano', 187.30, 'Meals'),
        ('2025-02-11', 'Marketing Campaign - Google Ads', 3200.00, 'Marketing'),
        ('2025-01-20', 'Software License - Adobe CC', 599.99, 'Software'),
        ('2025-03-05', 'Business Travel - Delta Airlines', 876.40, 'Travel'),
        ('2025-05-17', 'Conference Registration - TechSummit', 450.00, 'Training'),
        ('2025-02-28', 'Internet Service - AT&T', 189.99, 'Utilities'),
        ('2025-04-09', 'Printer Cartridges', 78.65, 'Office'),
        ('2025-01-15', 'Employee Training - Udemy Business', 299.00, 'Training'),
        ('2025-03-28', 'Catering - Q1 Review Meeting', 520.00, 'Meals'),
        ('2025-05-03', 'Cloud Hosting - AWS', 1920.00, 'IT'),
        ('2025-02-07', 'Office Furniture - Herman Miller Chair', 1245.00, 'Office'),
        ('2025-04-14', 'Marketing Campaign - LinkedIn Ads', 2750.00, 'Marketing'),
        ('2025-01-30', 'Software License - Slack', 240.00, 'Software'),
        ('2025-03-19', 'Business Travel - Marriott Hotel', 634.80, 'Travel'),
        ('2025-05-25', 'Team Building Event', 890.00, 'Entertainment'),
        ('2025-02-19', 'Legal Services - Contract Review', 1500.00, 'Legal'),
        ('2025-04-30', 'Office Supplies - Amazon Business', 156.75, 'Office'),
        ('2025-01-25', 'Video Conferencing - Zoom Pro', 199.90, 'Software'),
        ('2025-03-07', 'Shipping & Postage', 67.45, 'Logistics'),
        ('2025-05-12', 'Marketing Campaign - Facebook Ads', 1800.00, 'Marketing'),
        ('2025-02-14', 'Employee Appreciation Gifts', 425.00, 'Entertainment'),
        ('2025-04-03', 'Cloud Storage - Google Workspace', 350.00, 'IT'),
        ('2025-01-10', 'Accounting Software - QuickBooks', 499.00, 'Software'),
    ]

    data_border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row_idx, (date_str, desc, amount, category) in enumerate(transactions, 2):
        ws1.cell(row=row_idx, column=1, value=date_str)
        ws1.cell(row=row_idx, column=2, value=desc)
        cell_amount = ws1.cell(row=row_idx, column=3, value=amount)
        cell_amount.number_format = '$#,##0.00'
        ws1.cell(row=row_idx, column=4, value=category)
        for col in range(1, 5):
            ws1.cell(row=row_idx, column=col).border = data_border

    # Column widths
    ws1.column_dimensions['A'].width = 14
    ws1.column_dimensions['B'].width = 42
    ws1.column_dimensions['C'].width = 14
    ws1.column_dimensions['D'].width = 16

    # Freeze header row
    ws1.freeze_panes = 'A2'

    # --- Sheet 2: Empty (will hold the monthly summary after macro runs) ---
    ws2 = wb.create_sheet('Sheet2')
    # Sheet2 is intentionally empty — macro will populate it

    wb.save(XLSX_OUTPUT)
    print(f'Report automation workbook created: {XLSX_OUTPUT}')


def create_initial():
    create_macro_requirements_docx()
    create_report_automation_xlsx()

    # GUI-ready startup:
    # 1. Open report_automation.xlsx in LibreOffice Calc
    launch_gui(f'libreoffice --calc "{XLSX_OUTPUT}"', delay_sec=3.0)
    # 2. Open macro_requirements.docx in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DOCX_OUTPUT}"', delay_sec=2.0)

    print('GUI_READY: launched LibreOffice Calc and Writer with DISPLAY=:0')


create_initial()
