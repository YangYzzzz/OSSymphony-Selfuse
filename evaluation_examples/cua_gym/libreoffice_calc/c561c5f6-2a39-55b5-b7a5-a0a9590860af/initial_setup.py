"""
Initial Setup: Create assignment.docx with incorrectly formatted APA references
Task ID: osworld_multi_apps_misc_049
Domain: libreoffice_writer

Creates a student assignment document with a references section that has APA formatting errors:
- Authors listed with initials BEFORE surnames (wrong: J. Smith → correct: Smith, J.)
- Book and journal titles NOT italicized (should be italicized per APA 7th)
- URLs missing retrieval dates (should have "Retrieved [date] from" for changeable content)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
DESKTOP = f'{WORKDIR}/Desktop'
COURSEWORK_DIR = f'{DESKTOP}/coursework'
TASK_ID = 'osworld_multi_apps_misc_049'
OUTPUT = f'{COURSEWORK_DIR}/assignment.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure directories exist
    os.makedirs(COURSEWORK_DIR, exist_ok=True)

    doc = Document()

    # Set default margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("The Impact of Digital Learning Tools on Student Academic Performance")
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()

    # Author info
    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.add_run("Emily R. Patterson")

    dept_para = doc.add_paragraph()
    dept_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept_para.add_run("Department of Education, Westfield University")

    course_para = doc.add_paragraph()
    course_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    course_para.add_run("EDU 4820: Advanced Research Methods in Education")

    prof_para = doc.add_paragraph()
    prof_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    prof_para.add_run("Professor M. T. Harrison")

    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.add_run("November 15, 2024")

    doc.add_page_break()

    # Abstract
    abstract_heading = doc.add_heading("Abstract", level=1)
    abstract_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    abstract_body = doc.add_paragraph(
        "This study examines the effects of digital learning management systems (LMS) on undergraduate "
        "student academic performance across three semesters. Data was collected from 247 students at "
        "Westfield University using a mixed-methods approach combining quantitative grade analysis with "
        "qualitative survey responses. Results indicate that students who actively engaged with digital "
        "tools demonstrated a statistically significant improvement in their grade point average "
        "(M = 3.42, SD = 0.61) compared to those with minimal digital engagement (M = 2.87, SD = 0.74). "
        "Implications for instructional design and future research directions are discussed."
    )
    abstract_body.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    keywords_para = doc.add_paragraph()
    kw_label = keywords_para.add_run("Keywords: ")
    kw_label.italic = True
    keywords_para.add_run("digital learning, academic performance, LMS, higher education, instructional technology")
    keywords_para.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # Introduction
    doc.add_heading("Introduction", level=1)

    intro1 = doc.add_paragraph(
        "The proliferation of digital technologies in higher education has transformed the landscape "
        "of teaching and learning over the past two decades. Learning management systems such as Canvas, "
        "Blackboard, and Moodle have become standard infrastructure at universities worldwide, offering "
        "instructors and students a centralized platform for course materials, assignments, communication, "
        "and assessment."
    )
    intro1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    intro1.paragraph_format.first_line_indent = Inches(0.5)

    intro2 = doc.add_paragraph(
        "Despite widespread adoption, the relationship between digital tool utilization and academic "
        "outcomes remains contested in the literature. Some researchers argue that technology enhances "
        "engagement and facilitates personalized learning pathways, while others caution against "
        "over-reliance on digital platforms that may exacerbate existing inequalities in access and "
        "digital literacy."
    )
    intro2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    intro2.paragraph_format.first_line_indent = Inches(0.5)

    intro3 = doc.add_paragraph(
        "The present study seeks to contribute to this ongoing conversation by examining patterns of "
        "LMS engagement among undergraduate students and their relationship to academic performance "
        "indicators including GPA, assignment completion rates, and final examination scores."
    )
    intro3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    intro3.paragraph_format.first_line_indent = Inches(0.5)

    # Literature Review
    doc.add_heading("Literature Review", level=1)

    lit1 = doc.add_paragraph(
        "Research on educational technology has grown substantially since the early 2000s. "
        "Foundational work by B. S. Bloom (1984) on mastery learning provided a theoretical basis "
        "for personalized, self-paced instruction that digital tools have since operationalized at scale. "
        "Subsequent scholarship has investigated how specific digital affordances—such as immediate "
        "feedback, multimedia content, and collaborative tools—influence student motivation and retention."
    )
    lit1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    lit1.paragraph_format.first_line_indent = Inches(0.5)

    lit2 = doc.add_paragraph(
        "A meta-analysis by R. J. Means, Y. Toyama, R. Murphy, M. Bakia, and K. Jones (2009) "
        "evaluated 51 studies comparing online and face-to-face instruction, finding that blended "
        "approaches consistently produced superior learning outcomes. The study, titled "
        "Evaluation of Evidence-Based Practices in Online Learning, noted that the effect was particularly "
        "pronounced when online components incorporated active learning strategies."
    )
    lit2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    lit2.paragraph_format.first_line_indent = Inches(0.5)

    lit3 = doc.add_paragraph(
        "More recently, P. D. Dawson and P. Dawson (2017) examined LMS log data from 4,000 students "
        "across 14 universities in their book Student Success Prediction in MOOCs, revealing that "
        "frequency of platform logins was a stronger predictor of academic success than time spent "
        "on individual tasks. This finding underscores the importance of consistent engagement over "
        "intensive but sporadic use."
    )
    lit3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    lit3.paragraph_format.first_line_indent = Inches(0.5)

    # Methodology
    doc.add_heading("Methodology", level=1)

    meth1 = doc.add_paragraph(
        "This study employed a mixed-methods design integrating quantitative analysis of LMS usage "
        "logs with qualitative data from semi-structured interviews. Participants were 247 undergraduate "
        "students enrolled in six sections of core curriculum courses during the Fall 2023, Spring 2024, "
        "and Summer 2024 semesters at Westfield University."
    )
    meth1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    meth1.paragraph_format.first_line_indent = Inches(0.5)

    meth2 = doc.add_paragraph(
        "LMS activity data—including login frequency, time on task, resource downloads, and discussion "
        "forum participation—were extracted with institutional permission and linked to anonymized "
        "academic records. Qualitative interviews were conducted with a purposive subsample of 32 "
        "students representing varying levels of LMS engagement and academic performance."
    )
    meth2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    meth2.paragraph_format.first_line_indent = Inches(0.5)

    # Results
    doc.add_heading("Results", level=1)

    results1 = doc.add_paragraph(
        "Quantitative analysis revealed a moderate positive correlation between LMS login frequency "
        "and final course GPA (r = 0.47, p < .001). Students classified as high-engagement users "
        "(defined as logging in 4 or more times per week) achieved a mean GPA of 3.42 (SD = 0.61), "
        "while low-engagement users achieved a mean GPA of 2.87 (SD = 0.74)."
    )
    results1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    results1.paragraph_format.first_line_indent = Inches(0.5)

    results2 = doc.add_paragraph(
        "Qualitative themes identified from interview data included: (a) perceived convenience and "
        "accessibility of digital resources, (b) challenges related to self-regulation in online "
        "environments, and (c) the role of instructor responsiveness via digital platforms in "
        "sustaining student motivation."
    )
    results2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    results2.paragraph_format.first_line_indent = Inches(0.5)

    # Discussion
    doc.add_heading("Discussion", level=1)

    disc1 = doc.add_paragraph(
        "The findings support previous research suggesting that active digital engagement is associated "
        "with improved academic outcomes. However, this study's design does not permit causal inferences; "
        "students who are more motivated may engage more frequently with digital tools and also perform "
        "better academically for reasons unrelated to the tools themselves."
    )
    disc1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    disc1.paragraph_format.first_line_indent = Inches(0.5)

    disc2 = doc.add_paragraph(
        "Future research should investigate the mechanisms through which digital engagement influences "
        "learning, employing experimental or quasi-experimental designs that can better establish "
        "causality. Additionally, attention to equity implications—particularly for students with "
        "limited access to reliable internet or devices—is warranted."
    )
    disc2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    disc2.paragraph_format.first_line_indent = Inches(0.5)

    # Conclusion
    doc.add_heading("Conclusion", level=1)

    conc1 = doc.add_paragraph(
        "This study contributes empirical evidence to the ongoing conversation about digital learning "
        "in higher education. The moderate relationship between LMS engagement and academic performance "
        "suggests that instructors and institutions should actively promote consistent digital tool use, "
        "particularly among students who demonstrate low initial engagement."
    )
    conc1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    conc1.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # References section — INCORRECTLY FORMATTED (for the task to fix)
    references_heading = doc.add_heading("References", level=1)
    references_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Helper to add a reference with hanging indent
    def add_reference(doc, text_parts):
        """Add a reference paragraph with hanging indent.
        text_parts is a list of (text, italic) tuples.
        """
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        # Hanging indent: left indent 0.5in, first line indent -0.5in
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        para.paragraph_format.space_after = Pt(12)

        for text, italic in text_parts:
            run = para.add_run(text)
            if italic:
                run.italic = True
        return para

    # Reference 1: Journal article — WRONG: initials before surname
    # Correct APA 7: Bloom, B. S. (1984). The 2 sigma problem...
    # WRONG version: B. S. Bloom (1984). The 2 sigma problem...
    add_reference(doc, [
        ("B. S. Bloom (1984). The 2 sigma problem: The search for methods of group instruction as effective as one-to-one tutoring. ", False),
        ("Educational Researcher", False),  # WRONG: not italicized
        (", 13(6), 4–16.", False),
    ])

    # Reference 2: Government report with URL — WRONG: initials before surnames, no retrieval date
    # Correct APA 7: Means, B., Toyama, Y., Murphy, R., Bakia, M., & Jones, K. (2009). Evaluation of evidence-based practices in online learning...
    # WRONG: R. J. Means, Y. Toyama, R. Murphy, M. Bakia, & K. Jones (2009)...
    add_reference(doc, [
        ("R. J. Means, Y. Toyama, R. Murphy, M. Bakia, & K. Jones (2009). ", False),
        ("Evaluation of evidence-based practices in online learning: A meta-analysis and review of online learning studies", False),  # WRONG: not italicized
        (". U.S. Department of Education. https://www2.ed.gov/rschstat/eval/tech/evidence-based-practices/finalreport.pdf", False),
        # WRONG: no retrieval date for a URL
    ])

    # Reference 3: Book — WRONG: initials before surname, title not italicized
    # Correct APA 7: Dawson, P., & Dawson, S. L. (2017). Student success prediction in MOOCs...
    # WRONG: P. D. Dawson & P. Dawson (2017). Student success prediction in MOOCs...
    add_reference(doc, [
        ("P. D. Dawson & P. Dawson (2017). ", False),
        ("Student success prediction in MOOCs", False),  # WRONG: not italicized
        (". Springer.", False),
    ])

    # Reference 4: Journal article — WRONG: initials before surname, journal not italicized
    # Correct APA 7: Johnson, L., Adams Becker, S., Cummins, M., Estrada, V., Freeman, A., & Hall, C. (2016)...
    # WRONG: L. Johnson, S. Adams Becker, M. Cummins, V. Estrada, A. Freeman, & C. Hall (2016)...
    add_reference(doc, [
        ("L. Johnson, S. Adams Becker, M. Cummins, V. Estrada, A. Freeman, & C. Hall (2016). NMC horizon report: 2016 higher education edition. ", False),
        ("New Media Consortium", False),  # WRONG: not italicized (this is actually the publisher, but let's use a journal name here for the task)
        (".", False),
    ])

    # Reference 5: Website — WRONG: initials before surname, no retrieval date
    # Correct APA 7: Chen, B., deNoyelles, A., Patton, K., & Zydney, J. (2017)...
    # WRONG: B. Chen, A. deNoyelles, K. Patton, & J. Zydney (2017)...
    add_reference(doc, [
        ("B. Chen, A. deNoyelles, K. Patton, & J. Zydney (2017). Creating a community of inquiry in large-enrollment online courses: An exploratory study on the effect of course orientation. ", False),
        ("Online Learning", False),  # WRONG: not italicized
        (", 21(1). https://olj.onlinelearningconsortium.org/index.php/olj/article/view/986", False),
        # WRONG: no retrieval date for online journal URL
    ])

    # Reference 6: Book chapter — WRONG: initials before surname
    # Correct APA 7: Garrison, D. R., & Kanuka, H. (2004). Blended learning...
    # WRONG: D. R. Garrison & H. Kanuka (2004)...
    add_reference(doc, [
        ("D. R. Garrison & H. Kanuka (2004). Blended learning: Uncovering its transformative potential in higher education. ", False),
        ("The Internet and Higher Education", False),  # WRONG: not italicized
        (", 7(2), 95–105.", False),
    ])

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
