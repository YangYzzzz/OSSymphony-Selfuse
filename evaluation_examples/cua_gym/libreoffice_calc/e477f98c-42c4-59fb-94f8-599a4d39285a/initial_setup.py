"""
Initial Setup: LinkedList Tutorial with stub implementation
Task ID: osworld_multi_apps_misc_035
Domain: multi_apps (LibreOffice Writer + Python file editing)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_misc_035'
DESKTOP = f'{WORKDIR}/Desktop'
DOC_PATH = f'{DESKTOP}/LinkedList_Tutorial.docx'
PY_PATH = f'{DESKTOP}/linked_list.py'
RESULT_PATH = f'{DESKTOP}/ll_result.txt'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_tutorial_doc():
    """Create the LinkedList_Tutorial.docx explaining linked list operations."""
    doc = Document()

    # Title
    title = doc.add_heading("Singly Linked List - Tutorial", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction
    doc.add_heading("Introduction", level=1)
    intro = doc.add_paragraph(
        "A singly linked list is a linear data structure in which each element "
        "(called a node) contains a data field and a reference (link) to the next node "
        "in the sequence. The last node points to None, indicating the end of the list."
    )

    doc.add_paragraph(
        "Each node in a singly linked list contains:"
    )
    doc.add_paragraph("data - the value stored in the node", style="List Bullet")
    doc.add_paragraph("next - a pointer/reference to the next node (or None if last)", style="List Bullet")

    # Node Class
    doc.add_heading("The Node Class", level=1)
    doc.add_paragraph(
        "The Node class represents a single element in the linked list. "
        "It stores the data and a reference to the next node:"
    )

    node_code = doc.add_paragraph()
    node_run = node_code.add_run(
        "class Node:\n"
        "    def __init__(self, data):\n"
        "        self.data = data\n"
        "        self.next = None"
    )
    node_run.font.name = "Courier New"
    node_run.font.size = Pt(10)

    # LinkedList Class
    doc.add_heading("The LinkedList Class", level=1)
    doc.add_paragraph(
        "The LinkedList class manages a collection of nodes. It maintains a reference "
        "to the head (first) node of the list:"
    )

    ll_code = doc.add_paragraph()
    ll_run = ll_code.add_run(
        "class LinkedList:\n"
        "    def __init__(self):\n"
        "        self.head = None"
    )
    ll_run.font.name = "Courier New"
    ll_run.font.size = Pt(10)

    # Append Method
    doc.add_heading("Method: append(data)", level=1)
    doc.add_paragraph(
        "The append method adds a new node with the given data to the end of the list."
    )

    doc.add_heading("Algorithm:", level=2)
    doc.add_paragraph(
        "1. Create a new Node with the provided data.",
        style="List Number"
    )
    doc.add_paragraph(
        "2. If the list is empty (head is None), set head to the new node and return.",
        style="List Number"
    )
    doc.add_paragraph(
        "3. Otherwise, traverse the list starting from head until you reach the last node "
        "(the node whose next is None).",
        style="List Number"
    )
    doc.add_paragraph(
        "4. Set the next pointer of the last node to the new node.",
        style="List Number"
    )

    doc.add_heading("Implementation:", level=2)
    append_code = doc.add_paragraph()
    append_run = append_code.add_run(
        "def append(self, data):\n"
        "    new_node = Node(data)\n"
        "    if self.head is None:\n"
        "        self.head = new_node\n"
        "        return\n"
        "    current = self.head\n"
        "    while current.next is not None:\n"
        "        current = current.next\n"
        "    current.next = new_node"
    )
    append_run.font.name = "Courier New"
    append_run.font.size = Pt(10)

    # Delete Method
    doc.add_heading("Method: delete(data)", level=1)
    doc.add_paragraph(
        "The delete method removes the first node that contains the given data from the list."
    )

    doc.add_heading("Algorithm:", level=2)
    doc.add_paragraph(
        "1. If the list is empty (head is None), return immediately (nothing to delete).",
        style="List Number"
    )
    doc.add_paragraph(
        "2. If the head node contains the data to delete, update head to point to head.next and return.",
        style="List Number"
    )
    doc.add_paragraph(
        "3. Traverse the list keeping track of the current node and its previous node.",
        style="List Number"
    )
    doc.add_paragraph(
        "4. When you find a node whose data matches, set previous.next = current.next to bypass it.",
        style="List Number"
    )
    doc.add_paragraph(
        "5. If the data is not found, the list remains unchanged.",
        style="List Number"
    )

    doc.add_heading("Implementation:", level=2)
    delete_code = doc.add_paragraph()
    delete_run = delete_code.add_run(
        "def delete(self, data):\n"
        "    if self.head is None:\n"
        "        return\n"
        "    if self.head.data == data:\n"
        "        self.head = self.head.next\n"
        "        return\n"
        "    current = self.head\n"
        "    while current.next is not None:\n"
        "        if current.next.data == data:\n"
        "            current.next = current.next.next\n"
        "            return\n"
        "        current = current.next"
    )
    delete_run.font.name = "Courier New"
    delete_run.font.size = Pt(10)

    # Reverse Method
    doc.add_heading("Method: reverse()", level=1)
    doc.add_paragraph(
        "The reverse method reverses the order of nodes in the linked list in-place, "
        "so that the last node becomes the first."
    )

    doc.add_heading("Algorithm:", level=2)
    doc.add_paragraph(
        "1. Initialize three pointers: prev = None, current = head, next_node = None.",
        style="List Number"
    )
    doc.add_paragraph(
        "2. Iterate through the list: for each node, save next_node = current.next, "
        "then reverse current.next = prev.",
        style="List Number"
    )
    doc.add_paragraph(
        "3. Move prev to current, and current to next_node.",
        style="List Number"
    )
    doc.add_paragraph(
        "4. After the loop ends, set head = prev (the new first node).",
        style="List Number"
    )

    doc.add_heading("Implementation:", level=2)
    reverse_code = doc.add_paragraph()
    reverse_run = reverse_code.add_run(
        "def reverse(self):\n"
        "    prev = None\n"
        "    current = self.head\n"
        "    while current is not None:\n"
        "        next_node = current.next\n"
        "        current.next = prev\n"
        "        prev = current\n"
        "        current = next_node\n"
        "    self.head = prev"
    )
    reverse_run.font.name = "Courier New"
    reverse_run.font.size = Pt(10)

    # Display Method
    doc.add_heading("Helper Method: display()", level=1)
    doc.add_paragraph(
        "The display method prints all elements of the linked list in order, "
        "separated by ' -> ', followed by 'None' to indicate the end:"
    )

    display_code = doc.add_paragraph()
    display_run = display_code.add_run(
        "def display(self):\n"
        "    elements = []\n"
        "    current = self.head\n"
        "    while current is not None:\n"
        "        elements.append(str(current.data))\n"
        "        current = current.next\n"
        "    print(' -> '.join(elements) + ' -> None')"
    )
    display_run.font.name = "Courier New"
    display_run.font.size = Pt(10)

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Implement the three methods above in linked_list.py on the Desktop. "
        "The script already contains a Node class, a LinkedList class with stub methods, "
        "and test calls at the bottom. Use the implementations described in this tutorial."
    )

    doc.save(DOC_PATH)
    print(f'Tutorial document created: {DOC_PATH}')


def create_stub_python():
    """Create linked_list.py with stub methods (not implemented)."""
    stub_code = '''class Node:
    def __init__(self, data):
        self.data = data
        self.next = None


class LinkedList:
    def __init__(self):
        self.head = None

    def append(self, data):
        # TODO: Implement this method
        # Hint: Create a new Node and add it to the end of the list
        pass

    def delete(self, data):
        # TODO: Implement this method
        # Hint: Find and remove the first node with matching data
        pass

    def reverse(self):
        # TODO: Implement this method
        # Hint: Reverse the direction of all next pointers
        pass

    def display(self):
        elements = []
        current = self.head
        while current is not None:
            elements.append(str(current.data))
            current = current.next
        print(\' -> \'.join(elements) + \' -> None\')


# Test the LinkedList implementation
ll = LinkedList()

# Test append
ll.append(1)
ll.append(2)
ll.append(3)
ll.append(4)
ll.append(5)
print("After appending 1, 2, 3, 4, 5:")
ll.display()

# Test delete
ll.delete(3)
print("After deleting 3:")
ll.display()

# Test reverse
ll.reverse()
print("After reversing:")
ll.display()
'''

    with open(PY_PATH, 'w') as f:
        f.write(stub_code)
    print(f'Stub Python file created: {PY_PATH}')

    # Make sure ll_result.txt does NOT exist in initial state
    if os.path.exists(RESULT_PATH):
        os.remove(RESULT_PATH)
        print(f'Removed pre-existing result file: {RESULT_PATH}')


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    # Create tutorial document
    create_tutorial_doc()

    # Create stub Python file
    create_stub_python()

    # GUI-ready startup: open the tutorial in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DOC_PATH}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
