"""
Initial Setup: Cron Jobs Tutorial in LibreOffice Writer with existing crontab
Task ID: osworld_multi_apps_terminal_screenshot_011
Domain: multi_apps (libreoffice_writer + terminal + screenshot)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_terminal_screenshot_011'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DESKTOP = f'{WORKDIR}/Desktop'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def setup_crontab():
    """Set up a realistic crontab for the user with several scheduled entries."""
    crontab_content = (
        "# System maintenance cron jobs\n"
        "# m h  dom mon dow   command\n"
        "0 2 * * * /home/user/scripts/backup.sh >> /var/log/backup.log 2>&1\n"
        "*/15 * * * * /home/user/scripts/health_check.py\n"
        "0 8 * * 1-5 /home/user/scripts/send_report.sh\n"
        "30 23 * * 0 /home/user/scripts/weekly_cleanup.sh\n"
        "0 0 1 * * /home/user/scripts/monthly_stats.sh >> /var/log/stats.log\n"
    )
    # Write crontab by piping to crontab command
    proc = subprocess.Popen(
        ['crontab', '-'],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE
    )
    stdout, stderr = proc.communicate(crontab_content.encode())
    if proc.returncode != 0:
        print(f'Warning: crontab setup returned code {proc.returncode}: {stderr.decode()}')
    else:
        print('Crontab configured with scheduled entries.')


def create_tutorial_document():
    """Create a cron jobs tutorial document in LibreOffice Writer format."""
    doc = Document()

    # Title
    title = doc.add_heading('Linux Cron Jobs Tutorial', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction
    doc.add_heading('Introduction', level=1)
    intro = doc.add_paragraph(
        'Cron is a time-based job scheduler in Unix-like operating systems. '
        'It allows users to schedule scripts or commands to run automatically at specified times. '
        'The cron daemon reads the crontab (cron table) files and executes commands at the scheduled times.'
    )

    # Crontab syntax section
    doc.add_heading('Crontab Syntax', level=1)
    doc.add_paragraph(
        'The crontab file uses a specific syntax to define when and how often a command runs. '
        'Each line in a crontab represents a separate cron job.'
    )

    # Syntax explanation
    doc.add_heading('Time Field Format', level=2)
    doc.add_paragraph('A crontab entry consists of 5 time fields followed by the command:')

    # Add syntax as code-like paragraph
    syntax_para = doc.add_paragraph()
    run = syntax_para.add_run('* * * * *  command_to_execute')
    run.font.name = 'Courier New'
    run.font.size = Pt(11)

    doc.add_paragraph('Where the fields are:')
    doc.add_paragraph('Minute (0-59)', style='List Bullet')
    doc.add_paragraph('Hour (0-23)', style='List Bullet')
    doc.add_paragraph('Day of Month (1-31)', style='List Bullet')
    doc.add_paragraph('Month (1-12)', style='List Bullet')
    doc.add_paragraph('Day of Week (0-7, where 0 and 7 are Sunday)', style='List Bullet')

    # Special characters
    doc.add_heading('Special Characters', level=2)
    doc.add_paragraph(
        'Cron supports special characters to define complex schedules:'
    )
    doc.add_paragraph('* (asterisk) - matches all values', style='List Bullet')
    doc.add_paragraph(', (comma) - separates multiple values (e.g., 1,3,5)', style='List Bullet')
    doc.add_paragraph('- (dash) - defines a range (e.g., 1-5 for Mon-Fri)', style='List Bullet')
    doc.add_paragraph('/ (slash) - specifies step values (e.g., */15 every 15 minutes)', style='List Bullet')

    # Common examples
    doc.add_heading('Common Examples', level=1)
    doc.add_paragraph('Here are some common cron job examples:')

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Schedule'
    header_cells[1].text = 'Description'
    for cell in header_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    examples = [
        ('0 2 * * *', 'Run every day at 2:00 AM'),
        ('*/15 * * * *', 'Run every 15 minutes'),
        ('0 8 * * 1-5', 'Run at 8:00 AM on weekdays only'),
        ('30 23 * * 0', 'Run at 11:30 PM on Sundays'),
        ('0 0 1 * *', 'Run at midnight on the 1st of each month'),
    ]
    for i, (schedule, description) in enumerate(examples, 1):
        table.rows[i].cells[0].text = schedule
        table.rows[i].cells[1].text = description

    # Managing crontab section
    doc.add_heading('Managing Your Crontab', level=1)
    doc.add_paragraph(
        'You can manage your crontab using the crontab command with various flags:'
    )
    doc.add_paragraph('crontab -l  :  List all current cron jobs', style='List Bullet')
    doc.add_paragraph('crontab -e  :  Edit the crontab in your default editor', style='List Bullet')
    doc.add_paragraph('crontab -r  :  Remove all cron jobs', style='List Bullet')
    doc.add_paragraph('crontab -u username  :  Operate on another user\'s crontab (requires sudo)', style='List Bullet')

    # Viewing crontab section
    doc.add_heading('Viewing Your Current Crontab', level=1)
    doc.add_paragraph(
        'To view the cron jobs currently scheduled for your user account, '
        'open a terminal and run the following command:'
    )

    code_para = doc.add_paragraph()
    code_run = code_para.add_run('crontab -l')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(12)
    code_run.bold = True

    doc.add_paragraph(
        'This command will display all scheduled cron jobs for the current user. '
        'If no crontab has been set up, it will display "no crontab for <username>".'
    )

    # Logging section
    doc.add_heading('Logging and Monitoring', level=1)
    doc.add_paragraph(
        'By default, cron sends the output of jobs to the user\'s mail. '
        'To redirect output to a log file, append the desired path to the cron job:'
    )

    log_para = doc.add_paragraph()
    log_run = log_para.add_run('0 2 * * * /path/to/script.sh >> /var/log/myjob.log 2>&1')
    log_run.font.name = 'Courier New'
    log_run.font.size = Pt(10)

    doc.add_paragraph(
        'The >> appends output to the log file, and 2>&1 redirects error output to the same file.'
    )

    # Best practices
    doc.add_heading('Best Practices', level=1)
    doc.add_paragraph('Use absolute paths in your cron job commands and scripts.', style='List Number')
    doc.add_paragraph('Always redirect output to a log file for debugging purposes.', style='List Number')
    doc.add_paragraph('Test your scripts manually before adding them to crontab.', style='List Number')
    doc.add_paragraph('Add comments to your crontab entries to explain their purpose.', style='List Number')
    doc.add_paragraph('Use environment variables at the top of crontab if needed (e.g., MAILTO="").', style='List Number')

    doc.save(OUTPUT)
    print(f'Tutorial document created: {OUTPUT}')


def ensure_desktop():
    """Ensure Desktop directory exists."""
    os.makedirs(DESKTOP, exist_ok=True)
    print(f'Desktop directory ready: {DESKTOP}')


def create_initial():
    """Main setup function."""
    ensure_desktop()
    setup_crontab()
    create_tutorial_document()

    # GUI-ready startup: open LibreOffice Writer with the tutorial document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with cron jobs tutorial document (DISPLAY=:0)')


create_initial()
