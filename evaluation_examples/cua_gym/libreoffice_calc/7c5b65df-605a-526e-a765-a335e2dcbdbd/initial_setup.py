"""
Initial Setup: Research tools document with Chrome extensions list
Task ID: osworld_multi_apps_misc_008
Domain: multi_apps (LibreOffice Writer + Chrome)

Creates:
  - /home/user/Desktop/Research_tools.docx — document listing 5 Chrome extensions
  - Opens Chrome (no target extensions installed)
  - Opens LibreOffice Writer with the document
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_misc_008'
DESKTOP = f'{WORKDIR}/Desktop'
DOCX_PATH = f'{DESKTOP}/Research_tools.docx'

EXTENSIONS = [
    'Zotero Connector',
    'Google Scholar Button',
    'Unpaywall',
    'Research Rabbit',
    'Sci-Hub X Now',
]


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    # Create the Research_tools.docx document
    doc = Document()

    # Title
    title_para = doc.add_heading('Research Tools - Chrome Extensions', level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Introduction paragraph
    intro = doc.add_paragraph(
        'The following Chrome extensions are essential for academic research. '
        'Please install all of them in Chrome to enhance your research workflow.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # Extensions list heading
    list_heading = doc.add_paragraph('Chrome Extensions to Install:')
    list_heading.runs[0].bold = True
    list_heading.runs[0].font.size = Pt(12)

    # Add each extension as a bullet list item
    for ext_name in EXTENSIONS:
        bullet = doc.add_paragraph(ext_name, style='List Bullet')
        bullet.runs[0].font.size = Pt(11)

    # Additional notes paragraph
    doc.add_paragraph('')
    notes = doc.add_paragraph(
        'Note: All extensions listed above are available on the Chrome Web Store. '
        'They support citation management, open access discovery, and research discovery workflows.'
    )
    notes.paragraph_format.space_before = Pt(6)

    doc.save(DOCX_PATH)
    print(f'Initial document created: {DOCX_PATH}')

    # Kill any existing Chrome and LibreOffice instances to start fresh
    subprocess.run(['pkill', '-f', 'google-chrome'], capture_output=True)
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    time.sleep(2)

    # Launch Chrome (without extensions installed — initial state)
    launch_gui('google-chrome --new-window', delay_sec=3.0)

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{DOCX_PATH}"', delay_sec=2.0)

    print('GUI_READY: launched Chrome and LibreOffice Writer with DISPLAY=:0')


create_initial()
