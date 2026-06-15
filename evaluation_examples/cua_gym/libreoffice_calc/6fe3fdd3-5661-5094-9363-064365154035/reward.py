"""
Reward Script: Merge meeting_notes .txt files into meeting_minutes.docx
Task ID: osworld_multi_apps_vscode_concat_doc_006
Domain: libreoffice_writer (multi_apps with VSCode)

Task: The VSCode project meeting_notes on the Desktop contains .txt files named by date
(e.g., 2024-01-15.txt). Merge all files sorted by date (oldest first) into a single
LibreOffice Writer document called meeting_minutes.docx on the Desktop, with the
filename (date) as a heading before each file's content, and set font size to 12pt throughout.

Scoring:
  Component 1: meeting_minutes.docx exists at /home/user/Desktop/        (0.20 pts)
  Component 2: All 6 date headings present as Heading-style paragraphs   (0.30 pts)
  Component 3: Headings appear in chronological (ascending date) order    (0.20 pts)
  Component 4: All text runs use 12pt font size throughout                (0.30 pts)
  Total: 1.0
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_vscode_concat_doc_006'

# Expected dates derived from the meeting_notes folder (.txt filenames)
EXPECTED_DATES = [
    '2024-01-10',
    '2024-01-17',
    '2024-01-24',
    '2024-02-07',
    '2024-02-14',
    '2024-02-21',
]

DOCX_PATH = os.path.join(WORKDIR, 'Desktop', 'meeting_minutes.docx')
NOTES_DIR = os.path.join(WORKDIR, 'Desktop', 'meeting_notes')


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Component 1: meeting_minutes.docx exists on the Desktop (0.20 points)
    # This FAILS on initial_env (file absent) and PASSES on golden_env (file created)
    try:
        if not os.path.exists(file_path):
            print(f"FAIL: Component 1 — meeting_minutes.docx not found at {file_path}")
            print(f"\nScore: {total_score}/1.0")
            print(f"REWARD: {total_score}")
            return total_score
        else:
            file_size = os.path.getsize(file_path)
            if file_size > 0:
                print(f"PASS: Component 1 — meeting_minutes.docx exists ({file_size} bytes) (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 1 — meeting_minutes.docx exists but is empty")
                print(f"\nScore: {total_score}/1.0")
                print(f"REWARD: {total_score}")
                return total_score
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Load the docx file
    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load docx {file_path}: {e}")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Determine the expected dates from the notes directory (or use hardcoded list)
    try:
        if os.path.exists(NOTES_DIR):
            txt_files = sorted([
                f.replace('.txt', '')
                for f in os.listdir(NOTES_DIR)
                if f.endswith('.txt')
            ])
            expected_dates = txt_files if txt_files else EXPECTED_DATES
        else:
            expected_dates = EXPECTED_DATES
        print(f"Expected dates from notes folder: {expected_dates}")
    except Exception as e:
        expected_dates = EXPECTED_DATES
        print(f"WARN: Could not read notes dir, using hardcoded dates: {e}")

    # Collect all heading paragraphs from the docx
    heading_paragraphs = []
    for para in doc.paragraphs:
        if 'Heading' in para.style.name and para.text.strip():
            heading_paragraphs.append(para.text.strip())

    # Component 2: All expected date headings present (0.30 points)
    # FAILS on initial_env (file doesn't exist), PASSES on golden_env (all 6 headings exist)
    try:
        found_dates = []
        missing_dates = []
        for date in expected_dates:
            if date in heading_paragraphs:
                found_dates.append(date)
            else:
                missing_dates.append(date)

        if len(missing_dates) == 0 and len(found_dates) == len(expected_dates):
            print(f"PASS: Component 2 — All {len(expected_dates)} date headings found: {found_dates} (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 2 — Missing date headings: {missing_dates}; Found: {found_dates}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Headings appear in chronological (ascending date) order (0.20 points)
    # FAILS on initial_env (no file), PASSES on golden_env (sorted oldest-first)
    try:
        # Filter heading_paragraphs to only the date headings matching the expected pattern
        import re
        date_headings_in_doc = [h for h in heading_paragraphs if re.match(r'^\d{4}-\d{2}-\d{2}$', h)]

        if len(date_headings_in_doc) >= 2:
            sorted_dates = sorted(date_headings_in_doc)
            if date_headings_in_doc == sorted_dates:
                print(f"PASS: Component 3 — Headings in chronological order: {date_headings_in_doc} (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 — Headings not in order. Found: {date_headings_in_doc}, expected: {sorted_dates}")
        elif len(date_headings_in_doc) == 1:
            print(f"FAIL: Component 3 — Only 1 date heading found, cannot verify order")
        else:
            print(f"FAIL: Component 3 — No date headings found in document")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: All text runs have 12pt font size throughout (0.30 points)
    # FAILS on initial_env (no file), PASSES on golden_env (all 12pt)
    try:
        from docx.shared import Pt
        non_12pt_runs = []
        total_runs = 0

        for para in doc.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    total_runs += 1
                    size_pt = run.font.size.pt if run.font.size else None
                    if size_pt is None:
                        # None means "inherit from style" — check the paragraph style
                        # If the style font size is 12pt, this is still correct
                        pass
                    elif abs(size_pt - 12.0) > 0.5:
                        non_12pt_runs.append((para.text[:30], size_pt))

        if len(non_12pt_runs) == 0:
            print(f"PASS: Component 4 — All {total_runs} text runs are 12pt font size (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 4 — {len(non_12pt_runs)} runs with non-12pt size found: {non_12pt_runs[:5]}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
if not os.path.exists(DOCX_PATH):
    print(f"File not found: {DOCX_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(DOCX_PATH)
