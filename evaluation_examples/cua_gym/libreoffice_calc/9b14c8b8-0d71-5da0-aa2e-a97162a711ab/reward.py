"""
Reward Script: Process Q4 quarterly report preparation checklist
Task ID: osworld_multi_apps_doc_follow_instructions_011
Domain: libreoffice_calc + libreoffice_writer + os

Scoring Rubric (sum = 1.0):
  Component 1 (0.25): q4_master.ods exists with correct sheets and data
  Component 2 (0.20): q4_master.ods has 2 charts (line + pie) in Charts sheet
  Component 3 (0.30): q4_report.odt has placeholders replaced with real values
  Component 4 (0.15): q4_report.odt has Calibri branding applied to headings
  Component 5 (0.10): Q4_Report_2024.pdf exported to Desktop with bookmarks

All checks target task-introduced changes only. None of these pass on initial_env.
"""

import os
import re
import shutil
import zipfile

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_doc_follow_instructions_011'

# Expected paths
MASTER_ODS_PATH = os.path.join(WORKDIR, 'Documents', 'q4_master.ods')
REPORT_ODT_PATH = os.path.join(WORKDIR, 'Documents', 'q4_report.odt')
PDF_PATH = os.path.join(WORKDIR, 'Desktop', 'Q4_Report_2024.pdf')
PDF_ALT_PATH = os.path.join(WORKDIR, 'Desktop', 'q4_report.pdf')


def load_xlsx_via_copy(src_path):
    """
    q4_master.ods is internally xlsx-formatted (saved with .ods extension).
    Copy to a .xlsx temp file so openpyxl can load it.
    Returns (workbook, None) on success, (None, error_msg) on failure.
    """
    import openpyxl
    tmp = '/tmp/_reward_q4_master.xlsx'
    try:
        shutil.copy(src_path, tmp)
        wb = openpyxl.load_workbook(tmp)
        return wb, None
    except Exception as e:
        return None, str(e)


def verify_task():
    total_score = 0.0

    # -------------------------------------------------------------------------
    # Component 1: q4_master.ods exists with correct sheets and summary data
    #              (0.25 points)
    # Checks:
    #   - File exists at Documents/q4_master.ods
    #   - Has exactly 4 sheets: Sales Summary, Finance Summary, Ops Summary, Charts
    #   - Sales Summary has data rows (at least 4 data rows + header)
    #   - Sales Summary totals row contains numeric revenue value > 0
    # -------------------------------------------------------------------------
    try:
        if not os.path.exists(MASTER_ODS_PATH):
            print("FAIL: Component 1 — q4_master.ods not found in Documents/")
        else:
            wb, err = load_xlsx_via_copy(MASTER_ODS_PATH)
            if err or wb is None:
                print(f"FAIL: Component 1 — Cannot load q4_master.ods: {err}")
            else:
                expected_sheets = {'Sales Summary', 'Finance Summary', 'Ops Summary', 'Charts'}
                actual_sheets = set(wb.sheetnames)
                if not expected_sheets.issubset(actual_sheets):
                    missing = expected_sheets - actual_sheets
                    print(f"FAIL: Component 1 — Missing sheets: {missing}")
                else:
                    ws_sales = wb['Sales Summary']
                    # Check that data rows exist (rows 3-6 should have region data)
                    regions = []
                    for row_idx in range(3, 10):
                        cell_a = ws_sales.cell(row=row_idx, column=1).value
                        if cell_a and isinstance(cell_a, str) and cell_a.strip() and cell_a.upper() != 'TOTALS':
                            regions.append(cell_a)
                    # Check totals row has a numeric revenue value
                    totals_revenue = None
                    for row_idx in range(3, 15):
                        cell_a = ws_sales.cell(row=row_idx, column=1).value
                        if cell_a and isinstance(cell_a, str) and 'TOTAL' in cell_a.upper():
                            totals_revenue = ws_sales.cell(row=row_idx, column=4).value
                            break
                    if len(regions) >= 3 and totals_revenue and float(totals_revenue) > 0:
                        print(f"PASS: Component 1 — q4_master.ods has correct sheets and {len(regions)} data rows, "
                              f"total revenue={totals_revenue} (0.25 pts)")
                        total_score += 0.25
                    elif len(regions) < 3:
                        print(f"FAIL: Component 1 — Sales Summary has too few data rows: {len(regions)} (expected >= 3)")
                    else:
                        print(f"FAIL: Component 1 — Could not find totals revenue row in Sales Summary")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: q4_master.ods Charts sheet has 2 charts (line + pie)
    #              (0.20 points)
    # -------------------------------------------------------------------------
    try:
        if not os.path.exists(MASTER_ODS_PATH):
            print("FAIL: Component 2 — q4_master.ods not found")
        else:
            wb, err = load_xlsx_via_copy(MASTER_ODS_PATH)
            if err or wb is None:
                print(f"FAIL: Component 2 — Cannot load q4_master.ods: {err}")
            elif 'Charts' not in wb.sheetnames:
                print("FAIL: Component 2 — No 'Charts' sheet found")
            else:
                ws_charts = wb['Charts']
                charts = ws_charts._charts
                if len(charts) < 2:
                    print(f"FAIL: Component 2 — Expected 2 charts, found {len(charts)}")
                else:
                    # Check chart types
                    from openpyxl.chart import LineChart, PieChart
                    chart_types = [type(c).__name__ for c in charts]
                    has_line = any('Line' in ct for ct in chart_types)
                    has_pie = any('Pie' in ct for ct in chart_types)
                    if has_line and has_pie:
                        print(f"PASS: Component 2 — Charts sheet has {len(charts)} charts: "
                              f"{chart_types} (0.20 pts)")
                        total_score += 0.20
                    else:
                        print(f"FAIL: Component 2 — Expected line+pie charts, found: {chart_types}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: q4_report.odt placeholders replaced with actual values
    #              (0.30 points)
    # Checks that the 7 placeholder values in the table are replaced:
    #   - No {{...}} patterns remain
    #   - Table contains specific expected values
    # -------------------------------------------------------------------------
    try:
        if not os.path.exists(REPORT_ODT_PATH):
            print("FAIL: Component 3 — q4_report.odt not found")
        else:
            from docx import Document
            doc = Document(REPORT_ODT_PATH)
            # Collect all table cell text values
            placeholder_pattern = re.compile(r'\{\{[^}]+\}\}')
            all_table_text = []
            remaining_placeholders = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            all_table_text.append(text)
                            if placeholder_pattern.search(text):
                                remaining_placeholders.append(text)

            # Also check paragraphs
            for para in doc.paragraphs:
                if placeholder_pattern.search(para.text):
                    remaining_placeholders.append(para.text)

            if remaining_placeholders:
                print(f"FAIL: Component 3 — Unreplaced placeholders found: {remaining_placeholders[:5]}")
            else:
                # Check that specific expected values are present in the table
                expected_values = ['$8,587,100', '18,290', 'North America']
                found_values = []
                for tv in expected_values:
                    if any(tv in t for t in all_table_text):
                        found_values.append(tv)

                # Also check paragraph text for replaced values
                all_para_text = ' '.join(p.text for p in doc.paragraphs)
                for tv in expected_values:
                    if tv in all_para_text and tv not in found_values:
                        found_values.append(tv)

                if len(found_values) >= 2:
                    print(f"PASS: Component 3 — All placeholders replaced, "
                          f"found expected values: {found_values} (0.30 pts)")
                    total_score += 0.30
                else:
                    print(f"FAIL: Component 3 — Placeholders replaced but expected values not found. "
                          f"Found: {found_values}, Table content: {all_table_text[:5]}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: q4_report.odt has Calibri company branding on headings
    #              (0.15 points)
    # The golden file has Calibri font + color #1F3864 on heading runs.
    # Initial file has no explicit font set (None = inherited, no Calibri specified).
    # -------------------------------------------------------------------------
    try:
        if not os.path.exists(REPORT_ODT_PATH):
            print("FAIL: Component 4 — q4_report.odt not found")
        else:
            from docx import Document
            doc = Document(REPORT_ODT_PATH)
            calibri_headings = 0
            total_headings = 0
            for para in doc.paragraphs:
                # Check heading paragraphs (style contains Heading or para has h1 styling)
                style_name = para.style.name if para.style else ''
                if 'Heading' in style_name and para.text.strip():
                    total_headings += 1
                    # Check if any run has Calibri font explicitly set
                    for run in para.runs:
                        if run.font.name == 'Calibri':
                            calibri_headings += 1
                            break

            if total_headings == 0:
                print("FAIL: Component 4 — No heading paragraphs found")
            elif calibri_headings >= 2:
                print(f"PASS: Component 4 — {calibri_headings}/{total_headings} headings have Calibri branding (0.15 pts)")
                total_score += 0.15
            else:
                # Check body text for Calibri as a looser check
                calibri_body = sum(
                    1 for para in doc.paragraphs
                    for run in para.runs
                    if run.font.name == 'Calibri'
                )
                if calibri_body >= 5:
                    print(f"PASS: Component 4 — Calibri branding applied ({calibri_body} runs with Calibri font) (0.15 pts)")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 4 — Calibri branding not applied. "
                          f"{calibri_headings}/{total_headings} headings have Calibri, "
                          f"{calibri_body} body runs have Calibri")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -------------------------------------------------------------------------
    # Component 5: Q4_Report_2024.pdf exported to Desktop with bookmarks
    #              (0.10 points)
    # Checks:
    #   - PDF exists at ~/Desktop/Q4_Report_2024.pdf
    #   - PDF is valid (starts with %PDF)
    #   - PDF has /Outlines (bookmarks) with Count > 0
    # -------------------------------------------------------------------------
    try:
        # Accept either exact name or alternate
        pdf_path = None
        if os.path.exists(PDF_PATH):
            pdf_path = PDF_PATH
        elif os.path.exists(PDF_ALT_PATH):
            pdf_path = PDF_ALT_PATH

        if pdf_path is None:
            print(f"FAIL: Component 5 — PDF not found at {PDF_PATH} or {PDF_ALT_PATH}")
        else:
            with open(pdf_path, 'rb') as f:
                header = f.read(5)
                content = f.read()

            if not header.startswith(b'%PDF'):
                print(f"FAIL: Component 5 — File at {pdf_path} is not a valid PDF (header: {header})")
            else:
                # Check for bookmarks via /Outlines with Count > 0
                full_content = header + content
                has_outlines = b'/Outlines' in full_content
                # Find count value from outlines object
                outlines_count = 0
                count_matches = re.findall(rb'/Count\s+(\d+)', full_content)
                for cm in count_matches:
                    val = int(cm)
                    if val > 0:
                        outlines_count = val
                        break

                if has_outlines and outlines_count >= 3:
                    file_size_mb = (len(full_content)) / (1024 * 1024)
                    print(f"PASS: Component 5 — PDF exported to Desktop with {outlines_count} bookmarks, "
                          f"size={file_size_mb:.2f}MB (0.10 pts)")
                    total_score += 0.10
                elif has_outlines:
                    print(f"FAIL: Component 5 — PDF has /Outlines but bookmark count={outlines_count} (expected >= 3)")
                else:
                    print(f"FAIL: Component 5 — PDF exists but no /Outlines (bookmarks) found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -------------------------------------------------------------------------
    # Final Score
    # -------------------------------------------------------------------------
    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


if __name__ == '__main__':
    verify_task()
