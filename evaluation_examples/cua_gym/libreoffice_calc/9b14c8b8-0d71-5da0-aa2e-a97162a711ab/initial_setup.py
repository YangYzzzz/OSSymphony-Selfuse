"""
Initial Setup: Q4 Report Preparation Checklist with Multi-App Workflow
Task ID: osworld_multi_apps_doc_follow_instructions_011
Domain: libreoffice_calc (multi-app: calc + writer)

Creates:
  - /home/user/Documents/q4_prep_checklist.odt (13-item checklist)
  - /home/user/Documents/q4_sales.ods (Q4 sales data)
  - /home/user/Documents/q4_finance.ods (Q4 finance data)
  - /home/user/Documents/q4_ops.ods (Q4 operations data)
  - /home/user/Documents/q4_report.odt (report template with placeholders)
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
DOCS_DIR = '/home/user/Documents'
TASK_ID = 'osworld_multi_apps_doc_follow_instructions_011'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_docs_dir():
    os.makedirs(DOCS_DIR, exist_ok=True)


def create_checklist_odt():
    """Create q4_prep_checklist.odt with 13 checklist items using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    import lxml.etree as etree

    doc = Document()

    # Title
    title = doc.add_heading('Q4 2024 Report Preparation Checklist', level=1)
    title_run = title.runs[0] if title.runs else title.add_run('Q4 2024 Report Preparation Checklist')
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    intro = doc.add_paragraph(
        'This checklist documents all steps required to prepare the Q4 2024 quarterly report. '
        'Complete each item in sequence to ensure the final report is accurate and properly formatted.'
    )

    doc.add_paragraph('')

    # Section: Data Collection
    sec1 = doc.add_heading('Section 1: Data Collection', level=2)

    checklist_items = [
        (1,  'Collect summary row (Totals) from q4_sales.ods onto Sheet1 of a new file q4_master.ods. '
              'Include columns: Region, Product, Units Sold, Revenue, COGS, Gross Profit.'),
        (2,  'Collect summary row (Totals) from q4_finance.ods onto Sheet1 of q4_master.ods. '
              'Include columns: Category, Budget, Actual, Variance, Variance %.'),
        (3,  'Collect summary row (Totals) from q4_ops.ods onto Sheet1 of q4_master.ods. '
              'Include columns: Department, Headcount, Productivity Score, Efficiency Rating, Cost per Unit.'),
        (4,  'Review and validate all data entries in Sheet1 of q4_master.ods for accuracy and completeness. '
              'Ensure no blank cells remain in key data columns.'),
    ]

    for num, text in checklist_items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f'[ ] Task {num}: {text}')
        run.font.size = Pt(11)

    doc.add_paragraph('')

    # Section: Chart Creation
    sec2 = doc.add_heading('Section 2: Chart Creation', level=2)

    chart_items = [
        (5,  'Create a revenue trend line chart in q4_master.ods showing monthly revenue figures. '
              'Title the chart "Q4 2024 Revenue Trend". Place the chart on a sheet named "Charts".'),
        (6,  'Create an expense breakdown pie chart in q4_master.ods showing department expenses. '
              'Title the chart "Q4 2024 Expense Breakdown". Add to the "Charts" sheet.'),
    ]

    for num, text in chart_items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f'[ ] Task {num}: {text}')
        run.font.size = Pt(11)

    doc.add_paragraph('')

    # Section: Formatting
    sec3 = doc.add_heading('Section 3: Formatting & Styling', level=2)

    format_items = [
        (7,  'Apply executive summary formatting to q4_master.ods: bold header row with blue background '
              '(#4472C4), freeze the header row, set column widths to at least 15 characters. '
              'Add a title row "Q4 2024 Master Report" merged across all columns.'),
        (8,  'Apply conditional formatting to the Revenue column in q4_master.ods: highlight values '
              'above $500,000 in green (#00FF00) and values below $100,000 in red (#FF0000).'),
    ]

    for num, text in format_items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f'[ ] Task {num}: {text}')
        run.font.size = Pt(11)

    doc.add_paragraph('')

    # Section: Report Population
    sec4 = doc.add_heading('Section 4: Report Population', level=2)

    report_items = [
        (9,  'Open q4_report.odt and replace the placeholder {{TOTAL_REVENUE}} with the actual '
              'total revenue figure from q4_master.ods (sum of all regional revenues).'),
        (10, 'Replace all remaining placeholders in q4_report.odt: {{TOTAL_UNITS}}, {{GROSS_PROFIT}}, '
              '{{TOP_REGION}}, {{HEADCOUNT}}, {{EFFICIENCY_RATING}}, {{BUDGET_VARIANCE}}.'),
    ]

    for num, text in report_items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f'[ ] Task {num}: {text}')
        run.font.size = Pt(11)

    doc.add_paragraph('')

    # Section: Branding & Export
    sec5 = doc.add_heading('Section 5: Branding & Export', level=2)

    export_items = [
        (11, 'Apply company branding to q4_report.odt: change all body text font to "Calibri" size 11, '
              'headings to "Calibri" size 14 bold, and company color scheme (#1F3864 for headings, '
              '#2F5597 for subheadings).'),
        (12, 'Export q4_report.odt as a PDF file to the Desktop. Name the file "Q4_Report_2024.pdf". '
              'Ensure PDF bookmarks are created matching each H1 heading in the document.'),
        (13, 'Verify the exported PDF file Q4_Report_2024.pdf on the Desktop is email-ready: '
              'file size must be under 5MB. If larger, reduce image quality or compress images.'),
    ]

    for num, text in export_items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f'[ ] Task {num}: {text}')
        run.font.size = Pt(11)

    doc.add_paragraph('')

    # Footer note
    note = doc.add_paragraph(
        'Note: Complete all tasks in order. Each task builds on the previous. '
        'Mark each item by replacing "[ ]" with "[X]" when completed.'
    )
    note.runs[0].font.italic = True
    note.runs[0].font.size = Pt(10)

    output_path = f'{DOCS_DIR}/q4_prep_checklist.odt'
    doc.save(output_path)
    print(f'Created: {output_path}')


def create_sales_ods():
    """Create q4_sales.ods with realistic Q4 2024 sales data."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Q4 Sales Data'

    # Headers
    headers = ['Month', 'Region', 'Product', 'Units Sold', 'Unit Price', 'Revenue', 'COGS', 'Gross Profit']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFFFF")
        cell.fill = PatternFill(start_color="FF1F3864", end_color="FF1F3864", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")

    # Realistic Q4 data
    data = [
        # Month, Region, Product, Units Sold, Unit Price, Revenue, COGS, Gross Profit
        ['October',  'North America', 'Enterprise Suite',    1250, 850,  1062500, 425000, 637500],
        ['October',  'Europe',        'Professional Pack',   980,  620,  607600,  243040, 364560],
        ['October',  'Asia Pacific',  'Standard License',    2100, 340,  714000,  285600, 428400],
        ['October',  'Latin America', 'Basic Plan',          750,  180,  135000,   54000,  81000],
        ['November', 'North America', 'Enterprise Suite',    1380, 850,  1173000, 469200, 703800],
        ['November', 'Europe',        'Professional Pack',   1050, 620,  651000,  260400, 390600],
        ['November', 'Asia Pacific',  'Standard License',    2350, 340,  799000,  319600, 479400],
        ['November', 'Latin America', 'Basic Plan',          820,  180,  147600,   59040,  88560],
        ['December', 'North America', 'Enterprise Suite',    1620, 850,  1377000, 550800, 826200],
        ['December', 'Europe',        'Professional Pack',   1280, 620,  793600,  317440, 476160],
        ['December', 'Asia Pacific',  'Standard License',    2800, 340,  952000,  380800, 571200],
        ['December', 'Latin America', 'Basic Plan',          960,  180,  172800,   69120, 103680],
    ]

    for r, row_data in enumerate(data, 2):
        for c, val in enumerate(row_data, 1):
            ws.cell(row=r, column=c, value=val)
            if c in [5, 6, 7, 8]:  # monetary columns
                ws.cell(row=r, column=c).number_format = '$#,##0'

    # Totals row
    total_row = len(data) + 2
    ws.cell(row=total_row, column=1, value='TOTALS')
    ws.cell(row=total_row, column=2, value='All Regions')
    ws.cell(row=total_row, column=3, value='All Products')
    ws.cell(row=total_row, column=4, value=f'=SUM(D2:D{total_row-1})')
    ws.cell(row=total_row, column=5, value='N/A')
    ws.cell(row=total_row, column=6, value=f'=SUM(F2:F{total_row-1})')
    ws.cell(row=total_row, column=7, value=f'=SUM(G2:G{total_row-1})')
    ws.cell(row=total_row, column=8, value=f'=SUM(H2:H{total_row-1})')

    for c in [1, 2, 3]:
        ws.cell(row=total_row, column=c).font = Font(bold=True)
    for c in [4, 6, 7, 8]:
        ws.cell(row=total_row, column=c).font = Font(bold=True)
        ws.cell(row=total_row, column=c).number_format = '$#,##0'

    # Column widths
    col_widths = {'A': 12, 'B': 16, 'C': 20, 'D': 14, 'E': 12, 'F': 14, 'G': 14, 'H': 16}
    for col, width in col_widths.items():
        ws.column_dimensions[col].width = width

    output_path = f'{DOCS_DIR}/q4_sales.ods'
    wb.save(output_path)
    print(f'Created: {output_path}')


def create_finance_ods():
    """Create q4_finance.ods with realistic Q4 2024 finance data."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Q4 Finance Data'

    # Headers
    headers = ['Category', 'Month', 'Budget', 'Actual', 'Variance', 'Variance %']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFFFF")
        cell.fill = PatternFill(start_color="FF1F3864", end_color="FF1F3864", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")

    # Realistic financial data
    data = [
        ['Personnel',       'October',  2850000, 2823450, 26550,   0.0093],
        ['Infrastructure',  'October',   420000,  398750, 21250,   0.0506],
        ['Marketing',       'October',   650000,  712300, -62300, -0.0958],
        ['R&D',             'October',   380000,  375200,  4800,   0.0126],
        ['Administration',  'October',   185000,  179640,  5360,   0.0290],
        ['Personnel',       'November', 2850000, 2897320, -47320, -0.0166],
        ['Infrastructure',  'November',  420000,  403820,  16180,  0.0385],
        ['Marketing',       'November',  650000,  698540, -48540, -0.0747],
        ['R&D',             'November',  380000,  368900,  11100,  0.0292],
        ['Administration',  'November',  185000,  181250,   3750,  0.0203],
        ['Personnel',       'December', 2850000, 3012780, -162780, -0.0571],
        ['Infrastructure',  'December',  420000,  411300,   8700,   0.0207],
        ['Marketing',       'December',  650000,  785420, -135420, -0.2083],
        ['R&D',             'December',  380000,  371650,   8350,   0.0220],
        ['Administration',  'December',  185000,  193820,  -8820,  -0.0477],
    ]

    for r, row_data in enumerate(data, 2):
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val)
            if c in [3, 4, 5]:
                cell.number_format = '$#,##0'
            elif c == 6:
                cell.number_format = '0.00%'

    # Totals row
    total_row = len(data) + 2
    ws.cell(row=total_row, column=1, value='TOTALS')
    ws.cell(row=total_row, column=2, value='Q4 2024')
    for c in [3, 4, 5]:
        ws.cell(row=total_row, column=c, value=f'=SUM({chr(64+c)}2:{chr(64+c)}{total_row-1})')
        ws.cell(row=total_row, column=c).number_format = '$#,##0'
        ws.cell(row=total_row, column=c).font = Font(bold=True)
    ws.cell(row=total_row, column=6, value=f'=E{total_row}/C{total_row}')
    ws.cell(row=total_row, column=6).number_format = '0.00%'
    ws.cell(row=total_row, column=6).font = Font(bold=True)
    ws.cell(row=total_row, column=1).font = Font(bold=True)
    ws.cell(row=total_row, column=2).font = Font(bold=True)

    col_widths = {'A': 18, 'B': 12, 'C': 14, 'D': 14, 'E': 14, 'F': 12}
    for col, width in col_widths.items():
        ws.column_dimensions[col].width = width

    output_path = f'{DOCS_DIR}/q4_finance.ods'
    wb.save(output_path)
    print(f'Created: {output_path}')


def create_ops_ods():
    """Create q4_ops.ods with realistic Q4 2024 operations data."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Q4 Operations Data'

    # Headers
    headers = ['Department', 'Month', 'Headcount', 'Tickets Resolved', 'Productivity Score',
               'Efficiency Rating', 'Cost per Unit', 'SLA Compliance %']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFFFF")
        cell.fill = PatternFill(start_color="FF1F3864", end_color="FF1F3864", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")

    # Realistic operations data
    data = [
        ['Engineering',      'October',  142, 3845, 94.2, 'A', 48.50, 0.987],
        ['Customer Support', 'October',   85, 7234, 88.7, 'B', 22.30, 0.952],
        ['Sales',            'October',   67, 1250,  96.1, 'A', 85.40, 0.978],
        ['Marketing',        'October',   43,  320,  82.5, 'B', 62.10, 0.891],
        ['Operations',       'October',   38,  540,  91.3, 'A', 31.70, 0.943],
        ['Engineering',      'November', 145, 4012, 95.8, 'A', 47.20, 0.992],
        ['Customer Support', 'November',  87, 7891, 89.4, 'B', 21.80, 0.961],
        ['Sales',            'November',  69, 1380, 97.2, 'A', 83.90, 0.984],
        ['Marketing',        'November',  44,  298,  83.1, 'B', 63.40, 0.887],
        ['Operations',       'November',  39,  578,  92.4, 'A', 30.50, 0.955],
        ['Engineering',      'December', 148, 4320, 93.6, 'A', 50.10, 0.988],
        ['Customer Support', 'December',  90, 8456, 87.2, 'B', 23.10, 0.948],
        ['Sales',            'December',  72, 1620, 98.4, 'A', 81.20, 0.991],
        ['Marketing',        'December',  45,  412,  84.7, 'B', 61.80, 0.902],
        ['Operations',       'December',  40,  612,  93.8, 'A', 29.80, 0.967],
    ]

    for r, row_data in enumerate(data, 2):
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val)
            if c == 7:
                cell.number_format = '$#,##0.00'
            elif c == 8:
                cell.number_format = '0.0%'

    # Totals/Averages row
    total_row = len(data) + 2
    ws.cell(row=total_row, column=1, value='TOTALS/AVG')
    ws.cell(row=total_row, column=2, value='Q4 2024')
    ws.cell(row=total_row, column=3, value=f'=MAX(C2:C{total_row-1})')  # peak headcount
    ws.cell(row=total_row, column=4, value=f'=SUM(D2:D{total_row-1})')  # total tickets
    ws.cell(row=total_row, column=5, value=f'=AVERAGE(E2:E{total_row-1})')  # avg productivity
    ws.cell(row=total_row, column=5).number_format = '0.0'
    ws.cell(row=total_row, column=6, value='N/A')
    ws.cell(row=total_row, column=7, value=f'=AVERAGE(G2:G{total_row-1})')  # avg cost
    ws.cell(row=total_row, column=7).number_format = '$#,##0.00'
    ws.cell(row=total_row, column=8, value=f'=AVERAGE(H2:H{total_row-1})')  # avg SLA
    ws.cell(row=total_row, column=8).number_format = '0.0%'

    for c in range(1, 9):
        ws.cell(row=total_row, column=c).font = Font(bold=True)

    col_widths = {'A': 18, 'B': 12, 'C': 12, 'D': 18, 'E': 20, 'F': 18, 'G': 14, 'H': 18}
    for col, width in col_widths.items():
        ws.column_dimensions[col].width = width

    output_path = f'{DOCS_DIR}/q4_ops.ods'
    wb.save(output_path)
    print(f'Created: {output_path}')


def create_report_template_odt():
    """Create q4_report.odt with placeholders for the agent to fill."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Cover page
    cover_title = doc.add_heading('ACME Corporation', level=1)
    for run in cover_title.runs:
        run.font.size = Pt(24)
        run.font.bold = True

    subtitle = doc.add_paragraph('Q4 2024 Quarterly Business Report')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.bold = True

    date_p = doc.add_paragraph('Reporting Period: October – December 2024')
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in date_p.runs:
        run.font.size = Pt(12)
        run.font.italic = True

    doc.add_page_break()

    # Executive Summary
    h_exec = doc.add_heading('Executive Summary', level=1)

    exec_intro = doc.add_paragraph(
        'This report presents the consolidated financial and operational performance of '
        'ACME Corporation for Q4 2024. The following metrics represent company-wide aggregated data '
        'across all business units and geographic regions.'
    )

    doc.add_paragraph('')

    summary_table = doc.add_table(rows=8, cols=2)
    summary_table.style = 'Table Grid'

    metrics = [
        ('Metric', 'Value'),
        ('Total Revenue (Q4 2024)', '{{TOTAL_REVENUE}}'),
        ('Total Units Sold', '{{TOTAL_UNITS}}'),
        ('Gross Profit', '{{GROSS_PROFIT}}'),
        ('Top Performing Region', '{{TOP_REGION}}'),
        ('Total Headcount (Peak)', '{{HEADCOUNT}}'),
        ('Average Efficiency Rating', '{{EFFICIENCY_RATING}}'),
        ('Budget Variance', '{{BUDGET_VARIANCE}}'),
    ]

    for i, (label, value) in enumerate(metrics):
        cell_label = summary_table.cell(i, 0)
        cell_value = summary_table.cell(i, 1)
        cell_label.text = label
        cell_value.text = value

        if i == 0:
            for para in cell_label.paragraphs:
                for run in para.runs:
                    run.font.bold = True
            for para in cell_value.paragraphs:
                for run in para.runs:
                    run.font.bold = True

    doc.add_paragraph('')
    doc.add_page_break()

    # Sales Performance
    h_sales = doc.add_heading('Sales Performance', level=1)

    doc.add_paragraph(
        'Regional sales performance for Q4 2024 demonstrated strong growth across all territories. '
        'North America continued to lead in Enterprise Suite adoption, while Asia Pacific showed '
        'impressive volume in Standard License sales.'
    )

    doc.add_paragraph(
        'Total company revenue for Q4 2024 was {{TOTAL_REVENUE}}, with total units sold reaching '
        '{{TOTAL_UNITS}}. The gross profit margin remained healthy at {{GROSS_PROFIT}}.'
    )

    doc.add_paragraph(
        'The top performing region was {{TOP_REGION}}, contributing the highest revenue among all '
        'geographic territories.'
    )

    doc.add_paragraph('')
    doc.add_page_break()

    # Financial Summary
    h_finance = doc.add_heading('Financial Summary', level=1)

    doc.add_paragraph(
        'Q4 2024 financial performance showed disciplined cost management with an overall '
        'budget variance of {{BUDGET_VARIANCE}}. Major spending categories remained within '
        'acceptable ranges, with Marketing showing the highest variance due to an '
        'accelerated year-end campaign investment.'
    )

    doc.add_paragraph(
        'Full budget analysis is available in the attached q4_finance.ods supplementary file.'
    )

    doc.add_paragraph('')
    doc.add_page_break()

    # Operational Highlights
    h_ops = doc.add_heading('Operational Highlights', level=1)

    doc.add_paragraph(
        'Operational metrics for Q4 2024 reflect continued improvements in efficiency and '
        'service delivery. Peak headcount for the quarter was {{HEADCOUNT}} across all '
        'departments. The average efficiency rating was {{EFFICIENCY_RATING}}, reflecting '
        'strong performance across all operational units.'
    )

    doc.add_paragraph(
        'Engineering and Sales departments maintained the highest efficiency ratings (Grade A), '
        'while Customer Support handled record ticket volumes with strong SLA compliance.'
    )

    doc.add_paragraph('')
    doc.add_page_break()

    # Conclusion
    h_conclusion = doc.add_heading('Conclusion & Outlook', level=1)

    doc.add_paragraph(
        'Q4 2024 concluded with strong overall performance, meeting or exceeding key targets '
        'across revenue, profitability, and operational efficiency. The company is well-positioned '
        'for Q1 2025 with a healthy sales pipeline and streamlined operations.'
    )

    doc.add_paragraph(
        'Key priorities for Q1 2025 include expanding the Asia Pacific market presence, '
        'launching the next-generation Enterprise Suite platform, and maintaining cost discipline '
        'across all departments.'
    )

    output_path = f'{DOCS_DIR}/q4_report.odt'
    doc.save(output_path)
    print(f'Created: {output_path}')


def create_initial():
    print('Creating initial setup for osworld_multi_apps_doc_follow_instructions_011...')

    create_docs_dir()
    create_checklist_odt()
    create_sales_ods()
    create_finance_ods()
    create_ops_ods()
    create_report_template_odt()

    print('All initial files created.')
    print(f'Files in {DOCS_DIR}:')

    # GUI-ready startup: open the checklist in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DOCS_DIR}/q4_prep_checklist.odt"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with q4_prep_checklist.odt (DISPLAY=:0)')


create_initial()
