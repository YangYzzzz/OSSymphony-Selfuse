"""
Initial Setup: Research specialty coffee shops in 5 major cities and build comparison table + writer summary
Task ID: oswald_multi_apps_web_location_013
Domain: libreoffice_calc + libreoffice_writer (multi-app)
"""

import os
import shlex
import subprocess
import time
import shutil

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_web_location_013'
CALC_OUTPUT = f'{WORKDIR}/specialty_coffee_global.ods'
WRITER_OUTPUT = f'{WORKDIR}/Documents/coffee_city_guide.odt'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    import openpyxl
    from openpyxl.styles import Font, Alignment

    # Ensure Documents directory exists
    os.makedirs(f'{WORKDIR}/Documents', exist_ok=True)

    # Create an empty spreadsheet with just the column headers
    # The agent must fill in the actual research data
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Specialty Coffee Global"

    # Column headers only — agent must do the research and fill in data
    headers = [
        'City', 'Shop_Name', 'Address', 'Rating',
        'Specialty_Type', 'Hours', 'Roasts_In_House', 'Source_URL'
    ]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')

    # Set column widths for readability
    col_widths = {
        'A': 14,  # City
        'B': 28,  # Shop_Name
        'C': 35,  # Address
        'D': 8,   # Rating
        'E': 20,  # Specialty_Type
        'F': 20,  # Hours
        'G': 15,  # Roasts_In_House
        'H': 40,  # Source_URL
    }
    for col_letter, width in col_widths.items():
        ws.column_dimensions[col_letter].width = width

    # Save as xlsx first, then convert to ODS
    tmp_xlsx = f'{WORKDIR}/{TASK_ID}_tmp.xlsx'
    wb.save(tmp_xlsx)

    # Convert to proper ODS using LibreOffice headless
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'ods', tmp_xlsx, '--outdir', WORKDIR],
        capture_output=True, text=True, timeout=120
    )

    tmp_ods = f'{WORKDIR}/{TASK_ID}_tmp.ods'
    if os.path.exists(tmp_ods):
        shutil.move(tmp_ods, CALC_OUTPUT)
        os.remove(tmp_xlsx)
        print(f'Initial Calc file created (proper ODS): {CALC_OUTPUT}')
    else:
        # Fallback: use xlsx-format with .ods extension (LibreOffice can open it)
        shutil.move(tmp_xlsx, CALC_OUTPUT)
        print(f'Initial Calc file created (xlsx-as-ods fallback): {CALC_OUTPUT}')

    # Create an empty Writer document (placeholder)
    # The agent must write the city guide
    from docx import Document
    doc = Document()
    doc.add_heading('Global Specialty Coffee City Guide', level=0)
    p = doc.add_paragraph(
        'This document will contain a city-by-city guide to specialty coffee shops. '
        'Research each city using Chrome and fill in recommendations below.'
    )

    # Add placeholder sections for each city
    cities = ['New York', 'London', 'Tokyo', 'Melbourne', 'Seoul']
    for city in cities:
        doc.add_heading(city, level=1)
        doc.add_paragraph(
            f'[Research specialty coffee shops in {city} and add recommendations here.]'
        )

    # Save as .odt extension (docx format, LibreOffice can open it)
    tmp_docx = f'{WORKDIR}/{TASK_ID}_writer_tmp.docx'
    doc.save(tmp_docx)

    # Convert to proper ODT using LibreOffice headless
    result2 = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'odt', tmp_docx, '--outdir', f'{WORKDIR}/Documents/'],
        capture_output=True, text=True, timeout=120
    )
    tmp_odt = f'{WORKDIR}/Documents/{TASK_ID}_writer_tmp.odt'
    if os.path.exists(tmp_odt):
        shutil.move(tmp_odt, WRITER_OUTPUT)
        os.remove(tmp_docx)
        print(f'Initial Writer file created (proper ODT): {WRITER_OUTPUT}')
    else:
        shutil.move(tmp_docx, WRITER_OUTPUT)
        print(f'Initial Writer file created (docx-as-odt fallback): {WRITER_OUTPUT}')

    # GUI-ready startup: Open Chrome for research, then open the Calc file
    # Open Chrome browser for web research
    launch_gui('google-chrome --new-window "https://www.yelp.com/search?find_desc=specialty+coffee&find_loc=New+York"',
               delay_sec=2.0)

    # Open the initial Calc file for the agent
    launch_gui(f'libreoffice --calc "{CALC_OUTPUT}"', delay_sec=2.0)

    # Open the initial Writer file
    launch_gui(f'libreoffice --writer "{WRITER_OUTPUT}"', delay_sec=1.5)

    print('GUI_READY: launched Chrome, LibreOffice Calc, LibreOffice Writer with DISPLAY=:0')


create_initial()
