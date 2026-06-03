"""
Reward Script: Download PDF of first climate paper and document citations
Task ID: osworld_multi_apps_pdf_download_cite_005
Domain: multi_apps (LibreOffice Calc + Chrome + OS + LibreOffice Writer)
Scoring:
  Component 1 (0.40): climate_paper01.pdf exists and is a valid non-trivial PDF file
  Component 2 (0.40): climate_citations.docx exists and contains a citing paper title from the spreadsheet
  Component 3 (0.20): climate_citations.docx includes structured citation details (table or URL/DOI references)
Total: 1.0
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_pdf_download_cite_005'

# Titles of papers in the spreadsheet rows 2-5 (any could cite row 1)
SPREADSHEET_TITLES = [
    'Limiting Global Warming to 1.5°C: Implications for Carbon Budgets',
    'Climate change impacts under 1.5°C and 2°C of global warming: a focus on the IPCC SR1.5 report',
    'Net-zero emissions pathways: implications of the IPCC Special Report on Global Warming of 1.5°C',
    'Renewable energy and climate policy alignment in the post-Paris era',
]


def verify_task():
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # -----------------------------------------------------------------------
    # Component 1: climate_paper01.pdf exists and is a valid PDF (0.40 points)
    # This FAILS on initial_env (no PDF exists) and PASSES on golden_env.
    # We verify: file exists, is a valid PDF (header %PDF), and is non-trivially sized.
    # -----------------------------------------------------------------------
    pdf_path = os.path.join(WORKDIR, 'climate_paper01.pdf')
    try:
        if not os.path.isfile(pdf_path):
            print(f"FAIL: Component 1 — climate_paper01.pdf not found at {pdf_path}")
        else:
            file_size = os.path.getsize(pdf_path)
            # Verify it is a valid PDF (starts with %PDF header)
            with open(pdf_path, 'rb') as f:
                header = f.read(8)

            is_valid_pdf_header = header.startswith(b'%PDF')
            is_nontrivial_size = file_size >= 500

            if is_valid_pdf_header and is_nontrivial_size:
                print(f"PASS: Component 1 — climate_paper01.pdf is a valid PDF file "
                      f"(size={file_size} bytes, header={header[:8]!r}) (0.4 pts)")
                total_score += 0.4
            elif not is_valid_pdf_header:
                print(f"FAIL: Component 1 — {pdf_path} does not appear to be a valid PDF "
                      f"(header bytes: {header!r})")
            else:
                print(f"FAIL: Component 1 — climate_paper01.pdf exists but is suspiciously small "
                      f"({file_size} bytes), indicating it is not a real downloaded PDF")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: climate_citations.docx exists and contains a citing paper title (0.40 points)
    # The document must mention the title of at least one paper from the spreadsheet
    # that cites the IPCC report. This FAILS on initial_env (file doesn't exist).
    # -----------------------------------------------------------------------
    docx_path = os.path.join(WORKDIR, 'climate_citations.docx')
    try:
        if not os.path.isfile(docx_path):
            print(f"FAIL: Component 2 — climate_citations.docx not found at {docx_path}")
        else:
            from docx import Document
            doc = Document(docx_path)

            # Collect all text in the document (paragraphs + tables)
            all_text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text_parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            all_text_parts.append(cell.text.strip())

            full_doc_text_lower = ' '.join(all_text_parts).lower()

            # Check if any spreadsheet title appears in the document (exact or partial match)
            found_titles = []
            for title in SPREADSHEET_TITLES:
                if title.lower() in full_doc_text_lower:
                    found_titles.append(title)

            # Also check for partial matches using 6+ consecutive words from a title
            if not found_titles:
                for title in SPREADSHEET_TITLES:
                    words = title.lower().split()
                    if len(words) >= 6:
                        for i in range(len(words) - 5):
                            phrase = ' '.join(words[i:i+6])
                            if phrase in full_doc_text_lower:
                                found_titles.append(title)
                                break

            if found_titles:
                print(f"PASS: Component 2 — climate_citations.docx contains citing paper title(s): "
                      f"{found_titles} (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 2 — climate_citations.docx does not contain any known citing "
                      f"paper title from the spreadsheet. "
                      f"Document text snippet: {full_doc_text_lower[:400]!r}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: climate_citations.docx contains structured citation details (0.20 points)
    # The document should document the finding with structured detail — a table,
    # mention of DOI/URL, or year reference. This FAILS on initial_env.
    # -----------------------------------------------------------------------
    try:
        if not os.path.isfile(docx_path):
            print(f"FAIL: Component 3 — climate_citations.docx not found at {docx_path}")
        else:
            from docx import Document
            doc = Document(docx_path)

            has_table = len(doc.tables) > 0

            # Collect all text for URL/DOI detection
            all_text_lower = ' '.join(para.text for para in doc.paragraphs).lower()
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text_lower += ' ' + cell.text.lower()

            has_url_or_doi = (
                'ipcc.ch' in all_text_lower
                or 'doi.org' in all_text_lower
                or '10.1038' in all_text_lower
                or '10.1007' in all_text_lower
                or '10.1016' in all_text_lower
                or 'doi' in all_text_lower
            )
            has_year = any(str(y) in all_text_lower for y in [2018, 2019, 2020])

            if has_table and (has_url_or_doi or has_year):
                print(f"PASS: Component 3 — climate_citations.docx has structured citation details: "
                      f"table={has_table}, url/doi={has_url_or_doi}, year={has_year} (0.2 pts)")
                total_score += 0.2
            elif has_table:
                print(f"PASS: Component 3 — climate_citations.docx has a citation table (0.2 pts)")
                total_score += 0.2
            elif has_url_or_doi and has_year:
                print(f"PASS: Component 3 — climate_citations.docx has URL/DOI and year references (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — climate_citations.docx lacks structured citation details. "
                      f"has_table={has_table}, has_url_or_doi={has_url_or_doi}, has_year={has_year}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


verify_task()
