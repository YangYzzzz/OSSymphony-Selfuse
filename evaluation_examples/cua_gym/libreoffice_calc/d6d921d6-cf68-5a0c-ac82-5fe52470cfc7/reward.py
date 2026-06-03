"""
Reward Script: Download Shell Variables section from GNU Bash manual and save as bash_vars.docx
Task ID: osworld_multi_apps_web_to_doc_004
Domain: libreoffice_writer (multi_apps web_to_doc)
Scoring:
  Component 1: bash_vars.docx exists on Desktop and is a valid .docx file (0.3 pts)
  Component 2: Document contains the Shell Variables section title (0.3 pts)
  Component 3: Document contains key shell variable names from both sections (0.4 pts)
  Total: 1.0
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_web_to_doc_004'
FILE_PATH = '/home/user/Desktop/bash_vars.docx'


def verify_task(file_path):
    """
    Verify that bash_vars.docx exists on the Desktop and contains
    the Shell Variables section content from the GNU Bash manual.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Component 1: File exists on Desktop and is a valid .docx (0.3 points)
    # This fails on initial_env (no file) and passes on golden_env
    try:
        if not os.path.exists(file_path):
            print(f"FAIL: Component 1 — bash_vars.docx not found at {file_path}")
            print(f"\nScore: {total_score}/1.0")
            print(f"REWARD: {total_score}")
            return total_score

        # Try loading it as a valid docx
        from docx import Document
        try:
            doc = Document(file_path)
        except Exception as e:
            print(f"FAIL: Component 1 — File exists but is not a valid .docx: {e}")
            print(f"\nScore: {total_score}/1.0")
            print(f"REWARD: {total_score}")
            return total_score

        file_size = os.path.getsize(file_path)
        if file_size < 100:
            print(f"FAIL: Component 1 — File exists but is too small ({file_size} bytes) to contain real content")
            print(f"\nScore: {total_score}/1.0")
            print(f"REWARD: {total_score}")
            return total_score

        if file_size >= 100:
            print(f"PASS: Component 1 — bash_vars.docx exists on Desktop, valid .docx, size={file_size} bytes (0.3 pts)")
            total_score += 0.3

    except Exception as e:
        print(f"ERROR: Component 1 — {e}")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Re-load doc for further checks (already confirmed loadable)
    from docx import Document
    doc = Document(file_path)

    # Collect all text from the document
    all_paragraphs = [p.text.strip() for p in doc.paragraphs]
    full_text = '\n'.join(all_paragraphs)

    # Component 2: Document contains Shell Variables section title (0.3 points)
    # The section title should be present, indicating the correct section was downloaded
    try:
        # Check for "Shell Variables" in the title/headings
        heading_texts = [p.text.strip() for p in doc.paragraphs if 'Heading' in p.style.name]
        has_shell_vars_title = any(
            'shell variables' in h.lower() or 'shell variable' in h.lower()
            for h in heading_texts
        )

        # Also check in general text in case heading style wasn't preserved
        if not has_shell_vars_title:
            has_shell_vars_title = 'Shell Variables' in full_text or 'shell variables' in full_text.lower()

        if has_shell_vars_title:
            print(f"PASS: Component 2 — 'Shell Variables' section title found in document (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — 'Shell Variables' title not found. Headings: {heading_texts[:5]}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Document contains key shell variable names from both subsections (0.4 points)
    # Bourne Shell Variables: CDPATH, HOME, IFS, MAIL, MAILPATH, OPTARG, OPTIND, PATH, PS1, PS2
    # Bash Variables: BASH, BASHOPTS, BASHPID, BASH_ALIASES, BASH_ARGC, BASH_ARGV
    # We check for a meaningful subset across both groups
    try:
        bourne_vars = ['CDPATH', 'HOME', 'IFS', 'MAIL', 'PATH', 'PS1', 'PS2']
        bash_vars = ['BASH', 'BASHOPTS', 'BASHPID', 'BASH_ALIASES', 'BASH_ARGC']

        bourne_found = [v for v in bourne_vars if v in full_text]
        bash_found = [v for v in bash_vars if v in full_text]

        # Require at least 5 of 7 Bourne vars AND at least 3 of 5 Bash vars
        bourne_ok = len(bourne_found) >= 5
        bash_ok = len(bash_found) >= 3

        if bourne_ok and bash_ok:
            total_score += 0.4
            print(f"PASS: Component 3 — Key shell variables found: "
                  f"Bourne={bourne_found} ({len(bourne_found)}/7), "
                  f"Bash={bash_found} ({len(bash_found)}/5) (0.4 pts)")
        elif bourne_ok and not bash_ok:
            print(f"FAIL: Component 3 — Bourne vars OK ({len(bourne_found)}/7) but "
                  f"Bash vars insufficient: found {bash_found}, need >= 3 of {bash_vars}")
        elif not bourne_ok and bash_ok:
            print(f"FAIL: Component 3 — Bash vars OK ({len(bash_found)}/5) but "
                  f"Bourne vars insufficient: found {bourne_found}, need >= 5 of {bourne_vars}")
        else:
            print(f"FAIL: Component 3 — Insufficient variable coverage: "
                  f"Bourne={bourne_found} ({len(bourne_found)}/7), "
                  f"Bash={bash_found} ({len(bash_found)}/5)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Run verification against the canonical artifact path
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
