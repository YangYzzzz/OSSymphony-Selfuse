"""
Reward Script: Parse JSON config files, pretty-print, merge into LibreOffice Writer doc
Task ID: osworld_multi_apps_vscode_concat_doc_009
Domain: libreoffice_writer
Scoring:
  Component 1: configs_doc.docx exists on Desktop and has all 4 filenames as Heading-level sections (0.3 pts)
  Component 2: Body text paragraphs contain pretty-printed JSON content matching the source files (0.4 pts)
  Component 3: Body text (Normal paragraphs) have 11pt font size (0.3 pts)
"""

import os
import json

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_vscode_concat_doc_009'

# Expected JSON filenames sorted
EXPECTED_FILENAMES = sorted(['auth.json', 'cache.json', 'database.json', 'server.json'])
CONFIG_DIR = '/home/user/Desktop/config_files'
DOC_PATH = '/home/user/Desktop/configs_doc.docx'


def get_pretty_json(filename):
    """Load and pretty-print a JSON file from CONFIG_DIR."""
    path = os.path.join(CONFIG_DIR, filename)
    with open(path, 'r') as f:
        data = json.load(f)
    return json.dumps(data, indent=4)


def verify_task(doc_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition gate: doc must exist
    if not os.path.exists(doc_path):
        print(f"FAIL: configs_doc.docx not found at {doc_path}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load document {doc_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: 4 section headings matching JSON filenames (0.3 points)
    # Headings can be Heading 1 or any heading style
    try:
        heading_texts = []
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            if 'Heading' in style_name or style_name.startswith('heading'):
                heading_texts.append(para.text.strip())

        # Check that all 4 filenames appear as headings
        found_headings = [h for h in heading_texts if h in EXPECTED_FILENAMES]
        found_headings_set = set(found_headings)
        all_found = all(fname in found_headings_set for fname in EXPECTED_FILENAMES)

        if all_found:
            print(f"PASS: Component 1 — all 4 JSON filenames found as headings: {sorted(found_headings_set)} (0.3 pts)")
            total_score += 0.3
        else:
            missing = [f for f in EXPECTED_FILENAMES if f not in found_headings_set]
            print(f"FAIL: Component 1 — missing headings: {missing}. Found: {found_headings_set}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Body text paragraphs contain pretty-printed JSON content (0.4 points)
    # For each JSON file, its pretty-printed content should appear in the document body
    try:
        # Collect all non-heading paragraph text in the document
        body_texts = []
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            if 'Heading' not in style_name and para.text.strip():
                body_texts.append(para.text)

        full_body = '\n'.join(body_texts)

        # Check each JSON file's pretty-printed content appears in the document
        matched_count = 0
        for fname in EXPECTED_FILENAMES:
            try:
                expected_pretty = get_pretty_json(fname)
                # Normalize whitespace for comparison: strip and compare key tokens
                # Check that key parts of the pretty-printed JSON appear
                # Use first-level keys as identifiers
                data = json.loads(open(os.path.join(CONFIG_DIR, fname)).read())
                first_key = list(data.keys())[0]
                # Check that the opening brace with 4-space indent and the first key appear
                check_token = f'"  \\"{first_key}\\"'
                # More robust: check that the expected pretty-printed string appears
                # Allow for minor whitespace differences
                import re
                # Build a compact version to check key presence
                # Verify that the first key of the JSON appears in body with indent pattern
                indent_pattern = f'    "{first_key}"'
                if indent_pattern in full_body:
                    matched_count += 1
                    print(f"PASS: Component 2 — pretty-printed content for {fname} found in body")
                else:
                    print(f"FAIL: Component 2 — pretty-printed content for {fname} not found (expected indented key '{indent_pattern}')")
            except Exception as e2:
                print(f"ERROR: Component 2 — checking {fname}: {e2}")

        if matched_count == len(EXPECTED_FILENAMES):
            print(f"PASS: Component 2 — all {len(EXPECTED_FILENAMES)} files have pretty-printed JSON content (0.4 pts)")
            total_score += 0.4
        elif matched_count >= 2:
            partial = round(0.4 * matched_count / len(EXPECTED_FILENAMES), 2)
            print(f"PARTIAL: Component 2 — {matched_count}/{len(EXPECTED_FILENAMES)} files have pretty-printed content ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — only {matched_count}/{len(EXPECTED_FILENAMES)} files have pretty-printed content (0.0 pts)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Body text (Normal style paragraphs) have 11pt font size (0.3 points)
    try:
        normal_paras = [p for p in doc.paragraphs if p.style and 'Heading' not in p.style.name and p.text.strip()]
        if not normal_paras:
            print("FAIL: Component 3 — no Normal body text paragraphs found")
        else:
            correct_size_count = 0
            total_runs_checked = 0
            for para in normal_paras:
                for run in para.runs:
                    if run.text.strip():
                        total_runs_checked += 1
                        sz = run.font.size
                        sz_pt = sz.pt if sz else None
                        if sz_pt is not None and abs(sz_pt - 11.0) < 0.1:
                            correct_size_count += 1
                        else:
                            # If size is None it may be inherited - check style
                            if sz_pt is None:
                                # None means inherit from style; check the paragraph style size
                                style_sz = para.style.font.size if para.style else None
                                if style_sz and abs(style_sz.pt - 11.0) < 0.1:
                                    correct_size_count += 1

            if total_runs_checked == 0:
                print("FAIL: Component 3 — no runs with text found in body paragraphs")
            elif correct_size_count == total_runs_checked:
                print(f"PASS: Component 3 — all {total_runs_checked} body text runs have 11pt font size (0.3 pts)")
                total_score += 0.3
            elif correct_size_count > 0:
                ratio = correct_size_count / total_runs_checked
                if ratio >= 0.5:
                    print(f"PARTIAL: Component 3 — {correct_size_count}/{total_runs_checked} body runs have 11pt ({ratio:.1%}), awarding partial 0.15 pts")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 3 — only {correct_size_count}/{total_runs_checked} body runs have 11pt size")
            else:
                # Check: maybe all runs have None (inherited) which at the style level is 11pt
                # Try to verify via the first para's style
                first_normal = normal_paras[0]
                style_font_sz = None
                if first_normal.style and first_normal.style.font and first_normal.style.font.size:
                    style_font_sz = first_normal.style.font.size.pt
                print(f"FAIL: Component 3 — 0/{total_runs_checked} body runs have explicit 11pt size (style size: {style_font_sz} pt)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


if not os.path.exists(DOC_PATH):
    print(f"File not found: {DOC_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(DOC_PATH)
