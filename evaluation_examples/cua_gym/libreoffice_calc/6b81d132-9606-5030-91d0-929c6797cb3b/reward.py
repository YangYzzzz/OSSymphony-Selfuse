"""
Reward Script: Fix Works Cited section to follow MLA 9th edition format
Task ID: osworld_multi_apps_misc_044
Domain: libreoffice_writer (domain field in task config says libreoffice_calc, but actual task is Writer)
Scoring:
  - Component 1: Hanging indentation on all Works Cited reference entries (0.5 pts)
  - Component 2: MLA 9th edition date formatting with abbreviated months (0.3 pts)
  - Component 3: Edition abbreviation formatting (0.2 pts)
"""

import os
from docx import Document
from docx.shared import Inches

WORKDIR = '/home/user'
FILE_PATH = '/home/user/Desktop/manuscripts/journal_submission.docx'

# MLA 9th edition hanging indent: left indent = 0.5 inch (457200 EMU), first line indent = -0.5 inch (-457200 EMU)
MLA_HANGING_LEFT = 457200
MLA_HANGING_FIRST_LINE = -457200

# Works Cited entries start at paragraph index 18 (indices 18-24, 7 entries)
WORKS_CITED_START = 18
WORKS_CITED_END = 24  # inclusive


def verify_task(file_path):
    """
    Verify MLA 9th edition formatting in the Works Cited section.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print('CRITICAL: Cannot load file {}: {}'.format(file_path, e))
        print('REWARD: 0.0')
        return 0.0

    # Precondition: Check that Works Cited heading exists
    wc_heading_found = False
    for para in doc.paragraphs:
        if para.text.strip() == 'Works Cited':
            wc_heading_found = True
            break
    if not wc_heading_found:
        print('CRITICAL: Works Cited heading not found')
        print('REWARD: 0.0')
        return 0.0

    # Get all paragraphs starting from Works Cited heading
    all_paras = doc.paragraphs
    wc_start_idx = None
    for i, para in enumerate(all_paras):
        if para.text.strip() == 'Works Cited':
            wc_start_idx = i
            break

    if wc_start_idx is None:
        print('CRITICAL: Cannot locate Works Cited section index')
        print('REWARD: 0.0')
        return 0.0

    # Get all reference entry paragraphs (non-empty paragraphs after Works Cited heading)
    ref_paras = []
    for i in range(wc_start_idx + 1, len(all_paras)):
        para = all_paras[i]
        if para.text.strip():  # non-empty
            ref_paras.append((i, para))

    if len(ref_paras) == 0:
        print('CRITICAL: No reference entries found in Works Cited section')
        print('REWARD: 0.0')
        return 0.0

    print('Found {} Works Cited reference entries'.format(len(ref_paras)))

    # Component 1: Hanging indentation on ALL Works Cited reference entries (0.5 points)
    # MLA 9th edition requires hanging indent: left=0.5 inch, first_line=-0.5 inch
    try:
        hanging_correct_count = 0
        total_entries = len(ref_paras)
        for idx, para in ref_paras:
            pf = para.paragraph_format
            left = pf.left_indent
            fli = pf.first_line_indent
            has_correct_hanging = (
                left is not None and fli is not None and
                abs(left - MLA_HANGING_LEFT) < 100 and
                abs(fli - MLA_HANGING_FIRST_LINE) < 100
            )
            if has_correct_hanging:
                hanging_correct_count += 1
            else:
                print('FAIL: Entry [{}] missing hanging indent: left={}, fli={} | text={}'.format(
                    idx, left, fli, repr(para.text[:60])))

        if hanging_correct_count == total_entries:
            print('PASS: Component 1 — All {} entries have correct MLA hanging indentation (0.5 pts)'.format(total_entries))
            total_score += 0.5
        elif hanging_correct_count > 0:
            # Partial: give proportional credit
            partial = round(0.5 * hanging_correct_count / total_entries, 2)
            print('PARTIAL: Component 1 — {}/{} entries have correct hanging indent ({} pts)'.format(
                hanging_correct_count, total_entries, partial))
            total_score += partial
        else:
            print('FAIL: Component 1 — No entries have correct hanging indentation (0 pts)')
    except Exception as e:
        print('ERROR: Component 1 — {}'.format(e))

    # Component 2: MLA date formatting with abbreviated months (0.3 points)
    # MLA 9th edition uses abbreviated months: Jan., Feb., Mar., Apr., May, Jun., Jul., Aug., Sept., Oct., Nov., Dec.
    # And day-before-month order: "27 Mar. 2023" not "March 27, 2023"
    # Key checks (from actual data comparison):
    #   - "Jan. 2018" (not "January 2018") in Eubanks entry
    #   - "Sept. 2011" (not "September 2011") in Lankshear entry
    #   - "Sept. 2015" (not "September 2015") in Takayoshi entry
    #   - "27 Mar. 2023" and "15 Apr. 2023" format in OpenAI entry
    try:
        import re
        # Gather all text in works cited
        wc_text = ' '.join(para.text for _, para in ref_paras)

        date_checks = {
            'Jan. 2018': ('Jan. 2018', 'January 2018'),
            'Sept. 2011 or Sept. 2015': (['Sept. 2011', 'Sept. 2015'], ['September 2011', 'September 2015']),
            'day-before-month (e.g. 27 Mar.)': (['27 Mar.', '15 Apr.'], ['March 27', 'April 15']),
        }

        # Check: no full month names that should be abbreviated
        full_month_bad_patterns = [
            'January 2018', 'September 2011', 'September 2015',
            'March 27, 2023', 'April 15, 2023'
        ]
        mla_date_patterns = [
            'Jan. 2018', 'Sept. 2011', 'Sept. 2015',
            '27 Mar. 2023', '15 Apr. 2023'
        ]

        bad_dates_found = [p for p in full_month_bad_patterns if p in wc_text]
        good_dates_found = [p for p in mla_date_patterns if p in wc_text]

        if len(bad_dates_found) == 0 and len(good_dates_found) >= 3:
            print('PASS: Component 2 — MLA abbreviated date formats correct ({} good patterns found, 0.3 pts)'.format(len(good_dates_found)))
            total_score += 0.3
        elif len(bad_dates_found) == 0 and len(good_dates_found) >= 1:
            print('PARTIAL: Component 2 — Some MLA date formats correct ({} found, 0 bad): 0.15 pts'.format(len(good_dates_found)))
            total_score += 0.15
        else:
            print('FAIL: Component 2 — Found {} bad date formats: {}'.format(len(bad_dates_found), bad_dates_found))
            print('  Good patterns found: {}'.format(good_dates_found))
    except Exception as e:
        print('ERROR: Component 2 — {}'.format(e))

    # Component 3: Edition abbreviation formatting (0.2 points)
    # MLA 9th uses "3rd ed." not "Third Edition", "4th ed." not "Fourth Edition"
    try:
        wc_text_full = ' '.join(para.text for _, para in ref_paras)

        bad_edition_formats = ['Third Edition', 'Fourth Edition']
        good_edition_formats = ['3rd ed.', '4th ed.']

        bad_editions = [f for f in bad_edition_formats if f in wc_text_full]
        good_editions = [f for f in good_edition_formats if f in wc_text_full]

        if len(bad_editions) == 0 and len(good_editions) >= 2:
            print('PASS: Component 3 — Edition abbreviations correct: {} (0.2 pts)'.format(good_editions))
            total_score += 0.2
        elif len(bad_editions) == 0 and len(good_editions) == 1:
            print('PARTIAL: Component 3 — 1 of 2 edition abbreviations correct: {} (0.1 pts)'.format(good_editions))
            total_score += 0.1
        else:
            print('FAIL: Component 3 — Found bad edition formats: {}; good: {}'.format(bad_editions, good_editions))
    except Exception as e:
        print('ERROR: Component 3 — {}'.format(e))

    final_score = min(total_score, 1.0)
    print('\nScore: {}/1.0'.format(total_score))
    print('REWARD: {}'.format(final_score))
    return final_score


if not os.path.exists(FILE_PATH):
    print('File not found: {}'.format(FILE_PATH))
    print('REWARD: 0.0')
else:
    verify_task(FILE_PATH)
