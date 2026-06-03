"""
Initial Setup: journal_submission.docx with Works Cited section needing MLA 9th edition fixes
Task ID: osworld_multi_apps_misc_044
Domain: libreoffice_writer (multi_apps_misc)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_misc_044'
MANUSCRIPTS_DIR = f'{WORKDIR}/Desktop/manuscripts'
OUTPUT = f'{MANUSCRIPTS_DIR}/journal_submission.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure directory exists
    os.makedirs(MANUSCRIPTS_DIR, exist_ok=True)

    doc = Document()

    # Set margins (1 inch all around - standard for academic papers)
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title ---
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.line_spacing = 2.0
    run = title_para.add_run("Digital Literacy and Academic Discourse in the Age of Artificial Intelligence")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Author Info ---
    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.line_spacing = 2.0
    run = author_para.add_run("Elena M. Vasquez")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    affil_para = doc.add_paragraph()
    affil_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    affil_para.paragraph_format.line_spacing = 2.0
    run = affil_para.add_run("Department of English and Comparative Literature, Westbrook University")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Abstract heading ---
    abstract_heading = doc.add_paragraph()
    abstract_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    abstract_heading.paragraph_format.line_spacing = 2.0
    run = abstract_heading.add_run("Abstract")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Abstract body ---
    abstract_body = doc.add_paragraph()
    abstract_body.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    abstract_body.paragraph_format.line_spacing = 2.0
    abstract_body.paragraph_format.first_line_indent = Inches(0.5)
    run = abstract_body.add_run(
        "This paper examines the intersection of digital literacy and academic writing in the "
        "context of rapidly evolving artificial intelligence tools. As AI-assisted writing becomes "
        "increasingly prevalent in higher education, institutions face new challenges in defining "
        "authentic academic discourse. Drawing on theories of digital rhetoric and multimodal "
        "literacy, this study analyzes student perceptions of AI-generated content and its "
        "implications for scholarly communication. The findings suggest that faculty and students "
        "alike require updated frameworks for evaluating source credibility, argumentation, and "
        "intellectual originality in digitally mediated environments."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Introduction ---
    intro_heading = doc.add_paragraph()
    intro_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    intro_heading.paragraph_format.line_spacing = 2.0
    run = intro_heading.add_run("Introduction")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    intro_p1 = doc.add_paragraph()
    intro_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    intro_p1.paragraph_format.line_spacing = 2.0
    intro_p1.paragraph_format.first_line_indent = Inches(0.5)
    run = intro_p1.add_run(
        "The proliferation of large language models and AI writing assistants has fundamentally "
        "altered the landscape of academic composition. Scholars such as Selber and Takayoshi have "
        "long argued that digital literacy encompasses more than technical proficiency; it requires "
        "critical engagement with the epistemic assumptions embedded in technological systems. "
        "When students encounter AI tools capable of generating coherent, grammatically sophisticated "
        "prose, they must navigate complex questions about authorship, attribution, and intellectual "
        "labor."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    intro_p2 = doc.add_paragraph()
    intro_p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    intro_p2.paragraph_format.line_spacing = 2.0
    intro_p2.paragraph_format.first_line_indent = Inches(0.5)
    run = intro_p2.add_run(
        "Previous research on plagiarism detection and academic integrity has not adequately "
        "addressed the nuanced challenges posed by generative AI. Unlike traditional cases of "
        "plagiarism, AI-assisted writing blurs the boundary between assistance and appropriation. "
        "This study contributes to an emerging body of scholarship that seeks to articulate "
        "principled approaches to academic writing pedagogy in AI-saturated environments."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Literature Review ---
    lit_heading = doc.add_paragraph()
    lit_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    lit_heading.paragraph_format.line_spacing = 2.0
    run = lit_heading.add_run("Literature Review")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    lit_p1 = doc.add_paragraph()
    lit_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    lit_p1.paragraph_format.line_spacing = 2.0
    lit_p1.paragraph_format.first_line_indent = Inches(0.5)
    run = lit_p1.add_run(
        "Scholars in rhetoric and composition have explored digital literacies for more than two "
        "decades. Lankshear and Knobel define new literacies as socially recognized forms of "
        "generating and communicating meaning through digital tools. More recently, researchers "
        "have examined how algorithmic systems mediate meaning-making practices. Eubanks argues "
        "that automated systems often encode existing social inequities, a concern that extends "
        "to AI writing tools that reflect the biases embedded in their training data."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    lit_p2 = doc.add_paragraph()
    lit_p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    lit_p2.paragraph_format.line_spacing = 2.0
    lit_p2.paragraph_format.first_line_indent = Inches(0.5)
    run = lit_p2.add_run(
        "Haas and Flower's foundational work on reading to write demonstrates that skilled academic "
        "writers actively construct rhetorical contexts rather than passively absorbing information. "
        "This constructivist framework becomes especially relevant when considering how students "
        "interact with AI-generated content. Rather than simply consuming AI output, students must "
        "critically evaluate and rhetorically reframe such content to produce original scholarly work."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Methodology ---
    method_heading = doc.add_paragraph()
    method_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    method_heading.paragraph_format.line_spacing = 2.0
    run = method_heading.add_run("Methodology")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    method_p1 = doc.add_paragraph()
    method_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    method_p1.paragraph_format.line_spacing = 2.0
    method_p1.paragraph_format.first_line_indent = Inches(0.5)
    run = method_p1.add_run(
        "This study employed a mixed-methods approach combining survey instruments and discourse "
        "analysis. Undergraduate participants (n=247) from three institutions completed validated "
        "questionnaires assessing attitudes toward AI writing tools, academic integrity, and digital "
        "literacy self-efficacy. A subset of thirty participants submitted sample essays alongside "
        "AI-generated versions of the same prompts, which were then analyzed for rhetorical moves, "
        "evidence integration, and epistemic stance markers."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Discussion ---
    discussion_heading = doc.add_paragraph()
    discussion_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    discussion_heading.paragraph_format.line_spacing = 2.0
    run = discussion_heading.add_run("Discussion")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    discussion_p1 = doc.add_paragraph()
    discussion_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    discussion_p1.paragraph_format.line_spacing = 2.0
    discussion_p1.paragraph_format.first_line_indent = Inches(0.5)
    run = discussion_p1.add_run(
        "The survey results reveal significant ambivalence among students regarding the legitimacy "
        "of AI-assisted writing. While 68% of respondents reported using AI tools for brainstorming "
        "or outlining, only 23% considered AI-generated prose acceptable for direct inclusion in "
        "academic submissions. This discrepancy suggests that students possess intuitive notions of "
        "academic authorship even in the absence of explicit institutional guidelines."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Conclusion ---
    conclusion_heading = doc.add_paragraph()
    conclusion_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusion_heading.paragraph_format.line_spacing = 2.0
    run = conclusion_heading.add_run("Conclusion")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    conclusion_p1 = doc.add_paragraph()
    conclusion_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    conclusion_p1.paragraph_format.line_spacing = 2.0
    conclusion_p1.paragraph_format.first_line_indent = Inches(0.5)
    run = conclusion_p1.add_run(
        "As AI writing tools become ubiquitous, the academy must develop coherent frameworks for "
        "defining authentic scholarly contribution. This study suggests that students are ready to "
        "engage critically with AI-generated content provided they receive appropriate pedagogical "
        "scaffolding. Instructors should explicitly address the rhetorical dimensions of AI use, "
        "emphasizing the processes of critical evaluation and intellectual synthesis that distinguish "
        "academic writing from mere content generation."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Works Cited heading ---
    wc_heading = doc.add_paragraph()
    wc_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    wc_heading.paragraph_format.line_spacing = 2.0
    run = wc_heading.add_run("Works Cited")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Works Cited entries ---
    # INTENTIONALLY INCORRECT for MLA 9th edition:
    # 1. No hanging indent (incorrect - should have hanging indent)
    # 2. Missing publisher information
    # 3. Incorrect date formatting (e.g., "March 15, 2023" instead of "15 Mar. 2023")
    # 4. Wrong entry format overall

    entries = [
        # Entry 1: Missing publisher, bad date format, no hanging indent
        (
            'Eubanks, Virginia. Automating Inequality: How High-Tech Tools Profile, Police, and Punish the Poor. '
            'St. Martin\'s Press, January 2018.'
        ),
        # Entry 2: Missing publisher, wrong date format
        (
            'Haas, Christina, and Linda Flower. "Rhetorical Reading Strategies and the Construction of Meaning." '
            'College Composition and Communication, vol. 39, no. 2, May 1988, pp. 167-183.'
        ),
        # Entry 3: Missing publisher (journal missing volume info), bad date
        (
            'Lankshear, Colin, and Michele Knobel. New Literacies: Everyday Practices and Social Learning. '
            'Third Edition. Open University Press, September 2011.'
        ),
        # Entry 4: Missing access date for web source, wrong date format
        (
            'Selber, Stuart A. Multiliteracies for a Digital Age. '
            'Southern Illinois University Press, 2004.'
        ),
        # Entry 5: Missing publisher info, wrong date format for journal
        (
            'Takayoshi, Pamela. "Short-Form Writing: Studying Process in the Context of Contemporary Composing Technologies." '
            'Computers and Composition, vol. 37, September 2015, pp. 1-13.'
        ),
        # Entry 6: Website source missing publisher, wrong date format
        (
            'OpenAI. "GPT-4 Technical Report." OpenAI, March 27, 2023, https://openai.com/research/gpt-4. '
            'Accessed April 15, 2023.'
        ),
        # Entry 7: Missing publisher, wrong month abbreviation
        (
            'Wardle, Elizabeth, and Doug Downs. Writing About Writing: A College Reader. '
            'Fourth Edition. Bedford/St. Martin\'s, 2020.'
        ),
    ]

    for entry_text in entries:
        # INCORRECT: using regular paragraph indent instead of hanging indent
        entry_para = doc.add_paragraph()
        entry_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        entry_para.paragraph_format.line_spacing = 2.0
        # Wrong: setting first_line_indent positively (regular indent, not hanging)
        entry_para.paragraph_format.left_indent = Inches(0)
        entry_para.paragraph_format.first_line_indent = Inches(0)
        run = entry_para.add_run(entry_text)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Also create some other files in the manuscripts folder (realistic environment)
    # Create a draft notes file
    notes_path = f'{MANUSCRIPTS_DIR}/submission_notes.txt'
    with open(notes_path, 'w') as f:
        f.write("Submission Notes - Journal of Digital Rhetoric\n")
        f.write("Target word count: 6000-8000 words\n")
        f.write("Current word count: approx. 4200\n")
        f.write("Deadline: April 30, 2024\n")
        f.write("Reviewers: 2 double-blind reviewers\n")
        f.write("TODO: Fix Works Cited formatting before submission\n")

    print(f'Notes file created: {notes_path}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
