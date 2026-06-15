"""
Initial Setup: Create Security_extensions.docx on Desktop with Chrome extension list
Task ID: osworld_multi_apps_misc_010
Domain: multi_apps (LibreOffice Writer + Chrome)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_multi_apps_misc_010'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/Security_extensions.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    # Create Security_extensions.docx
    doc = Document()

    # Title
    title = doc.add_heading('Recommended Security & Privacy Chrome Extensions', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph
    intro = doc.add_paragraph(
        'The following Chrome extensions have been reviewed and approved by the IT Security '
        'Department. Please install all of these extensions to ensure your browser meets '
        'our organization\'s security and privacy standards.'
    )

    doc.add_paragraph('')  # blank line

    # Section header
    doc.add_heading('Required Extensions', level=2)

    # Extension list
    extensions = [
        ('uBlock Origin',
         'A widely used, efficient ad blocker and content filter. Blocks ads, '
         'trackers, and malware domains to improve browsing security.'),
        ('Privacy Badger',
         'Developed by the Electronic Frontier Foundation (EFF). Automatically '
         'learns to block invisible trackers based on their behavior.'),
        ('HTTPS Everywhere',
         'Encrypts your communications with many major websites, making your '
         'browsing more secure. Developed by EFF and the Tor Project.'),
        ('DuckDuckGo Privacy Essentials',
         'Provides private search, tracker blocking, and site encryption. '
         'Includes a Privacy Grade rating for each website you visit.'),
        ('ClearURLs',
         'Removes tracking elements from URLs automatically to protect your '
         'privacy while browsing and sharing links.'),
    ]

    for i, (name, description) in enumerate(extensions, 1):
        # Extension number and name
        p = doc.add_paragraph()
        run_num = p.add_run(f'{i}. ')
        run_num.bold = True
        run_name = p.add_run(name)
        run_name.bold = True
        run_name.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Dark blue

        # Description
        desc_para = doc.add_paragraph(f'   {description}')
        desc_para.paragraph_format.space_after = Pt(6)

    doc.add_paragraph('')  # blank line

    # Footer instructions
    instructions = doc.add_paragraph(
        'Installation Instructions: Open Google Chrome, navigate to the Chrome Web Store '
        '(chrome.google.com/webstore), search for each extension by name, and click '
        '"Add to Chrome" for each one listed above.'
    )
    instructions.paragraph_format.space_before = Pt(12)
    run_bold = instructions.runs[0]

    # Note paragraph
    note = doc.add_paragraph(
        'Note: If you experience any issues during installation, please contact '
        'the IT Help Desk at helpdesk@company.com or call extension 4357.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Kill any existing Chrome instances before launching
    subprocess.run(['pkill', '-f', 'google-chrome'], capture_output=True)
    time.sleep(2)

    # GUI-ready startup: open Chrome first, then LibreOffice Writer with the docx
    # Launch Chrome with remote debugging port
    launch_gui(
        'google-chrome --remote-debugging-port=1337 --no-first-run --no-default-browser-check',
        delay_sec=3.0
    )

    # Launch LibreOffice Writer with the security extensions document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)

    print('GUI_READY: launched Chrome and LibreOffice Writer with DISPLAY=:0')


create_initial()
