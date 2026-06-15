"""
Initial Setup: Research paper with IEEE-format violations in references section
Task ID: osworld_multi_apps_misc_042
Domain: libreoffice_writer

Creates /home/user/Desktop/submissions/research_paper.docx with:
- A research paper structure (title, abstract, introduction, methodology, results, conclusion)
- A references section at the end containing IEEE-format violations:
  * Missing brackets around reference numbers (e.g., "1." instead of "[1]")
  * Missing publisher info for some entries
  * Inconsistent ordering (not cited in document order)
  * Some entries lack required fields (volume, pages, year)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_misc_042'
SUBMISSIONS_DIR = f'{WORKDIR}/Desktop/submissions'
OUTPUT = f'{SUBMISSIONS_DIR}/research_paper.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Create submissions directory if it doesn't exist
    os.makedirs(SUBMISSIONS_DIR, exist_ok=True)

    doc = Document()

    # --- Title ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run(
        "Deep Learning Approaches for Autonomous Vehicle Perception: "
        "A Comprehensive Survey"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)

    # --- Authors ---
    authors_para = doc.add_paragraph()
    authors_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    authors_para.add_run(
        "Michael Zhang1, Sarah Hoffmann2, Raj Patel3, Elena Voronova1"
    )

    affil_para = doc.add_paragraph()
    affil_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    affil_run = affil_para.add_run(
        "1Stanford University, 2TU Munich, 3IIT Delhi"
    )
    affil_run.font.size = Pt(10)

    doc.add_paragraph()  # blank line

    # --- Abstract ---
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Autonomous vehicles rely on robust perception systems to navigate complex "
        "real-world environments. This survey provides a comprehensive review of deep "
        "learning methodologies applied to vehicle perception tasks, including object "
        "detection, semantic segmentation, depth estimation, and sensor fusion. We "
        "analyze over 150 recent publications and categorize approaches by architecture "
        "type, dataset benchmarks, and real-world performance metrics. Our findings "
        "indicate that transformer-based models have surpassed traditional convolutional "
        "approaches on most standard benchmarks, while multi-modal fusion techniques "
        "demonstrate superior robustness under adverse weather conditions. We conclude "
        "with an assessment of open research challenges and directions for future work."
    )

    # --- Keywords ---
    kw_para = doc.add_paragraph()
    kw_run = kw_para.add_run("Keywords: ")
    kw_run.bold = True
    kw_para.add_run(
        "autonomous vehicles, deep learning, object detection, semantic segmentation, "
        "sensor fusion, LiDAR, camera, transformer"
    )

    # --- 1. Introduction ---
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "The development of autonomous vehicles (AVs) has accelerated dramatically "
        "over the past decade, driven by advances in sensing technology, computational "
        "hardware, and machine learning algorithms [1]. Perception systems form the "
        "critical foundation of any AV pipeline, responsible for transforming raw "
        "sensor data from cameras, LiDAR, and radar into structured scene representations "
        "that downstream planning and control modules can act upon [2]."
    )
    doc.add_paragraph(
        "Early AV perception systems relied heavily on hand-crafted features and "
        "rule-based classifiers. The introduction of deep convolutional neural networks "
        "(CNNs) by Krizhevsky et al. [3] revolutionized computer vision and subsequently "
        "transformed the AV perception landscape. More recently, vision transformers "
        "have demonstrated competitive or superior performance on standard benchmarks [4], "
        "prompting a wave of transformer-based architectures for 3D object detection, "
        "BEV representation learning, and occupancy prediction [5]."
    )
    doc.add_paragraph(
        "This survey aims to provide researchers and practitioners with a structured "
        "overview of the field. Section 2 describes our survey methodology. Section 3 "
        "covers camera-based perception. Section 4 addresses LiDAR-based methods. "
        "Section 5 discusses multi-modal fusion. Section 6 surveys depth estimation, "
        "and Section 7 concludes with open challenges."
    )

    # --- 2. Methodology ---
    doc.add_heading("2. Survey Methodology", level=1)
    doc.add_paragraph(
        "We conducted a systematic literature search using Google Scholar, IEEE Xplore, "
        "arXiv, and ACM Digital Library. Search terms included combinations of "
        "'autonomous driving', 'perception', 'object detection', 'LiDAR', '3D detection', "
        "'semantic segmentation', and 'sensor fusion'. Papers published between 2018 and "
        "2024 were considered, with some seminal earlier works included for context."
    )
    doc.add_paragraph(
        "Inclusion criteria required that papers: (1) propose a novel deep learning "
        "method for AV perception, (2) evaluate on a publicly available benchmark "
        "dataset, and (3) report quantitative performance metrics. A total of 163 "
        "papers met these criteria and are discussed throughout this survey."
    )

    # --- 3. Camera-Based Perception ---
    doc.add_heading("3. Camera-Based Perception", level=1)
    doc.add_paragraph(
        "Camera-based perception encompasses 2D and 3D object detection, lane detection, "
        "traffic sign recognition, and semantic/panoptic segmentation. The KITTI dataset "
        "established early benchmarks for 3D detection from monocular images [6], while "
        "nuScenes provides a more challenging multi-camera 360-degree evaluation [7]."
    )
    doc.add_paragraph(
        "YOLO-family detectors have achieved real-time performance on 2D detection tasks, "
        "with YOLOv8 reporting 53.9% mAP on COCO at 80 FPS on an NVIDIA A100 GPU [8]. "
        "For 3D detection, BEVFormer leverages multi-camera features projected into a "
        "bird's-eye-view representation using deformable attention, achieving 56.9% NDS "
        "on the nuScenes test set [9]."
    )

    # --- 4. LiDAR-Based Methods ---
    doc.add_heading("4. LiDAR-Based Perception", level=1)
    doc.add_paragraph(
        "LiDAR sensors provide accurate 3D geometric information but at higher cost and "
        "with lower texture resolution than cameras. PointNet was among the first "
        "architectures to directly process raw point clouds for classification and "
        "segmentation [10]. VoxelNet discretizes point clouds into voxels and applies "
        "3D convolutions, trading some accuracy for computational efficiency [11]."
    )
    doc.add_paragraph(
        "More recent sparse convolutional methods, exemplified by CenterPoint and "
        "VoxelNeXt, have achieved state-of-the-art results on the Waymo Open Dataset "
        "benchmark. CenterPoint reports 65.5 mAPH on the Waymo 3D detection leaderboard [12]."
    )

    # --- 5. Sensor Fusion ---
    doc.add_heading("5. Multi-Modal Sensor Fusion", level=1)
    doc.add_paragraph(
        "Fusing camera and LiDAR data can overcome the limitations of each modality. "
        "Early-fusion approaches combine raw sensor inputs before feature extraction, "
        "while late-fusion methods merge detection outputs from independently processed "
        "streams. Intermediate-fusion architectures, such as PointPainting and "
        "TransFusion, project LiDAR points onto image feature maps to enrich point "
        "features with semantic context [13][14]."
    )

    # --- 6. Depth Estimation ---
    doc.add_heading("6. Monocular Depth Estimation", level=1)
    doc.add_paragraph(
        "Estimating metric depth from a single image is an ill-posed problem, yet "
        "recent self-supervised methods have achieved remarkable results. Monodepth2 "
        "trains jointly on monocular video sequences using photometric reprojection "
        "loss and outperforms fully supervised methods on some metrics [15]. Foundation "
        "models such as Depth Anything show strong zero-shot generalization to diverse "
        "scenes [16]."
    )

    # --- 7. Conclusion ---
    doc.add_heading("7. Conclusion", level=1)
    doc.add_paragraph(
        "This survey has reviewed the state of the art in deep learning for autonomous "
        "vehicle perception. Transformer-based architectures now dominate top leaderboard "
        "positions, and multi-modal fusion is essential for robust all-weather performance. "
        "Key open challenges include: (1) long-tail scenario generalization, (2) "
        "real-time inference at scale, (3) interpretability and uncertainty quantification, "
        "and (4) domain adaptation across geographic regions and sensor configurations. "
        "We anticipate that the integration of large vision-language models will further "
        "reshape the field in coming years."
    )

    # --- References (with IEEE violations) ---
    # Violations intentionally included:
    # - Using "1." notation instead of "[1]"
    # - Missing publisher for book entries
    # - Some entries out of expected numeric order
    # - Missing volume/issue/pages for some journal entries
    # - Inconsistent punctuation (some use semicolons, some missing commas)
    doc.add_heading("References", level=1)

    references_bad = [
        # 1. Missing brackets, uses period notation
        "1. L. Fridman, \"Deep Learning for Autonomous Vehicles,\" "
        "IEEE Transactions on Intelligent Vehicles vol. 3 no. 2 pp. 98-107, 2018.",

        # 2. Missing publisher info for book
        "2. S. Thrun, W. Burgard, and D. Fox, Probabilistic Robotics, 2005.",

        # 3. Correct except uses period instead of brackets
        "3. A. Krizhevsky, I. Sutskever, and G. E. Hinton, \"ImageNet Classification "
        "with Deep Convolutional Neural Networks,\" in Advances in Neural Information "
        "Processing Systems, vol. 25, pp. 1097-1105, 2012.",

        # 4. Missing conference location and year separated incorrectly
        "4. A. Dosovitskiy et al. \"An Image is Worth 16x16 Words: Transformers for "
        "Image Recognition at Scale\" in ICLR 2021",

        # 5. Out of order (should be [5] but content references item appearing earlier)
        "5. J. Hu, L. Shen, and G. Sun, \"Squeeze-and-Excitation Networks,\" "
        "in Proc. CVPR, pp. 7132-7141 2018.",

        # 6. Missing commas and incomplete citation
        "6. A. Geiger P. Lenz and R. Urtasun \"Are we ready for Autonomous Driving? "
        "The KITTI Vision Benchmark Suite\" CVPR 2012.",

        # 7. Missing volume/issue info
        "7. H. Caesar et al., \"nuScenes: A Multimodal Dataset for Autonomous Driving,\" "
        "in Proc. CVPR, 2020.",

        # 8. Uses semicolon instead of comma, wrong bracket style
        "8. G. Jocher et al.; YOLOv8 by Ultralytics; 2023. [Online] Available: "
        "https://github.com/ultralytics/ultralytics",

        # 9. Missing required field (no volume/pages)
        "9. Z. Liu et al., \"BEVFormer: Learning Bird's-Eye-View Representation from "
        "Multi-Camera Images via Spatiotemporal Transformers,\" ECCV 2022.",

        # 10. Partially correct but missing period at end and missing page numbers
        "10. C. R. Qi, H. Su, K. Mo, and L. J. Guibas, \"PointNet: Deep Learning on "
        "Point Sets for 3D Classification and Segmentation,\" in Proc. CVPR, 2017",

        # 11. Missing comma between author names
        "11. Y. Zhou and O. Tuzel \"VoxelNet: End-to-End Learning for Point Cloud Based "
        "3D Object Detection\" in Proc. CVPR pp. 4490-4499, 2018.",

        # 12. Uses wrong citation style mixing APA and IEEE
        "12. Yin, T., Zhou, X., & Krähenbühl, P. (2021). Center-based 3D Object "
        "Detection and Tracking. CVPR, 11784-11793.",

        # 13. Missing pages and incomplete venue
        "13. S. Vora, A. H. Lang, B. Helou, and O. Beijbom, \"PointPainting: Sequential "
        "Fusion for 3D Object Detection,\" CVPR 2020.",

        # 14. Uses et al. without comma before it, missing journal info
        "14. X. Bai et al \"TransFusion: Robust LiDAR-Camera Fusion for 3D Object "
        "Detection with Transformers\" CVPR 2022.",

        # 15. Missing issue number, inconsistent formatting
        "15. C. Godard, O. Mac Aodha, M. Firman and G. J. Brostow, \"Digging Into "
        "Self-Supervised Monocular Depth Estimation,\" in Proc. ICCV, pp. 3828-3838 2019.",

        # 16. Missing conference proceedings label
        "16. L. Yang et al., \"Depth Anything: Unleashing the Power of Large-Scale "
        "Unlabeled Data,\" CVPR, 2024.",
    ]

    for ref in references_bad:
        doc.add_paragraph(ref)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Add a few other files in submissions directory to make it realistic
    other_doc = Document()
    other_doc.add_paragraph("Supplementary materials for submission.")
    other_doc.save(f'{SUBMISSIONS_DIR}/supplementary.docx')

    cover_doc = Document()
    cover_doc.add_paragraph("Cover letter for journal submission.")
    cover_doc.save(f'{SUBMISSIONS_DIR}/cover_letter.docx')

    # GUI-ready startup: open the research paper in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with research_paper.docx on DISPLAY=:0')


create_initial()
