"""
Initial Setup: Create data_report.odt with embedded JSON array of 20 statistical records
Task ID: osworld_multi_apps_media_doc_edit_010
Domain: libreoffice_writer (multi-app)
"""

import os
import shlex
import subprocess
import time
import json

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_media_doc_edit_010'
DOCUMENTS_DIR = f'{WORKDIR}/documents'
OUTPUT = f'{DOCUMENTS_DIR}/data_report.odt'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Create documents directory if it doesn't exist
    os.makedirs(DOCUMENTS_DIR, exist_ok=True)

    # Define 20 statistical records covering Jan-Dec 2024 with 2 metrics
    # Metrics: "CPU_Usage" (%) and "Memory_Usage" (GB)
    # 10 records per metric, spread across months
    records = [
        {"date": "2024-01-15", "metric_name": "CPU_Usage",    "value": 42.3,  "unit": "%"},
        {"date": "2024-01-15", "metric_name": "Memory_Usage", "value": 6.1,   "unit": "GB"},
        {"date": "2024-02-15", "metric_name": "CPU_Usage",    "value": 55.8,  "unit": "%"},
        {"date": "2024-02-15", "metric_name": "Memory_Usage", "value": 7.4,   "unit": "GB"},
        {"date": "2024-03-15", "metric_name": "CPU_Usage",    "value": 38.2,  "unit": "%"},
        {"date": "2024-03-15", "metric_name": "Memory_Usage", "value": 5.9,   "unit": "GB"},
        {"date": "2024-04-15", "metric_name": "CPU_Usage",    "value": 61.4,  "unit": "%"},
        {"date": "2024-04-15", "metric_name": "Memory_Usage", "value": 8.2,   "unit": "GB"},
        {"date": "2024-05-15", "metric_name": "CPU_Usage",    "value": 47.9,  "unit": "%"},
        {"date": "2024-05-15", "metric_name": "Memory_Usage", "value": 6.8,   "unit": "GB"},
        {"date": "2024-06-15", "metric_name": "CPU_Usage",    "value": 73.1,  "unit": "%"},
        {"date": "2024-06-15", "metric_name": "Memory_Usage", "value": 9.5,   "unit": "GB"},
        {"date": "2024-07-15", "metric_name": "CPU_Usage",    "value": 68.5,  "unit": "%"},
        {"date": "2024-07-15", "metric_name": "Memory_Usage", "value": 10.1,  "unit": "GB"},
        {"date": "2024-08-15", "metric_name": "CPU_Usage",    "value": 52.3,  "unit": "%"},
        {"date": "2024-08-15", "metric_name": "Memory_Usage", "value": 7.7,   "unit": "GB"},
        {"date": "2024-09-15", "metric_name": "CPU_Usage",    "value": 44.6,  "unit": "%"},
        {"date": "2024-09-15", "metric_name": "Memory_Usage", "value": 6.5,   "unit": "GB"},
        {"date": "2024-10-15", "metric_name": "CPU_Usage",    "value": 59.7,  "unit": "%"},
        {"date": "2024-10-15", "metric_name": "Memory_Usage", "value": 8.9,   "unit": "GB"},
    ]

    json_str = json.dumps(records, indent=2)

    # Write ODT using python-docx (saved as .docx then converted, or use odfpy)
    # We'll use python-docx and save as .docx, then use libreoffice to convert to .odt
    # OR we can create a proper ODT using odfpy
    # For simplicity and reliability, let's create it as a proper docx first, then convert

    # Create using python-docx
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Title
    title = doc.add_heading('Server Performance Data Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph
    intro = doc.add_paragraph(
        'This report contains raw server performance metrics collected throughout 2024. '
        'The data below is provided in JSON format for programmatic analysis. '
        'It includes CPU usage and memory usage measurements recorded monthly from '
        'January 2024 through October 2024.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # Section heading
    doc.add_heading('Raw Data (JSON Format)', level=1)

    # Description paragraph
    desc = doc.add_paragraph(
        'The following JSON array contains 20 statistical records. Each record includes: '
        'date (YYYY-MM-DD), metric_name (CPU_Usage or Memory_Usage), value (numeric), '
        'and unit (% or GB).'
    )
    desc.paragraph_format.space_after = Pt(6)

    # JSON data block
    json_para = doc.add_paragraph()
    run = json_para.add_run(json_str)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    json_para.paragraph_format.space_before = Pt(6)
    json_para.paragraph_format.space_after = Pt(12)

    # Additional section with notes
    doc.add_heading('Data Collection Notes', level=1)
    notes = doc.add_paragraph(
        'Metrics were collected from the primary production server cluster. '
        'CPU_Usage values represent the average utilization percentage across all cores. '
        'Memory_Usage values represent total RAM consumed in gigabytes. '
        'All measurements were taken at 15:00 UTC on the 15th of each month.'
    )
    notes.paragraph_format.space_after = Pt(12)

    # Contact info
    doc.add_heading('Contact', level=2)
    doc.add_paragraph('For questions about this report, contact: sysadmin@company.internal')

    # Save as docx first in temp location
    temp_docx = f'{DOCUMENTS_DIR}/data_report_temp.docx'
    doc.save(temp_docx)
    print(f'Temporary docx created: {temp_docx}')

    # Convert to ODT using LibreOffice
    convert_env = os.environ.copy()
    convert_env["DISPLAY"] = ":0"
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'odt',
         '--outdir', DOCUMENTS_DIR, temp_docx],
        capture_output=True, text=True, env=convert_env, timeout=60
    )
    print(f'LibreOffice conversion stdout: {result.stdout}')
    print(f'LibreOffice conversion stderr: {result.stderr}')

    # The converted file will be named data_report_temp.odt - rename it
    converted_file = f'{DOCUMENTS_DIR}/data_report_temp.odt'
    if os.path.exists(converted_file):
        os.rename(converted_file, OUTPUT)
        print(f'Renamed to: {OUTPUT}')
    elif os.path.exists(OUTPUT):
        print(f'Output already exists at: {OUTPUT}')
    else:
        # Fallback: create ODT directly using odfpy
        print('Conversion failed, creating ODT directly with odfpy...')
        create_odt_direct(json_str)

    # Clean up temp file
    if os.path.exists(temp_docx):
        os.remove(temp_docx)

    print(f'Initial file created: {OUTPUT}')

    # Verify file exists
    if os.path.exists(OUTPUT):
        size = os.path.getsize(OUTPUT)
        print(f'File size: {size} bytes')
    else:
        print(f'ERROR: Output file not found at {OUTPUT}')

    # GUI-ready startup: open data_report.odt in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with data_report.odt on DISPLAY=:0')


def create_odt_direct(json_str):
    """Fallback: create ODT directly using odfpy."""
    try:
        from odf.opendocument import OpenDocumentText
        from odf.style import Style, TextProperties, ParagraphProperties
        from odf.text import H, P, Span
        from odf import teletype

        doc = OpenDocumentText()

        # Title
        h1 = H(outlinelevel=1)
        h1.addText('Server Performance Data Report')
        doc.text.addElement(h1)

        # Intro
        p1 = P()
        p1.addText(
            'This report contains raw server performance metrics collected throughout 2024. '
            'The data below is provided in JSON format for programmatic analysis.'
        )
        doc.text.addElement(p1)

        # Section heading
        h2 = H(outlinelevel=2)
        h2.addText('Raw Data (JSON Format)')
        doc.text.addElement(h2)

        # JSON data
        p_json = P()
        p_json.addText(json_str)
        doc.text.addElement(p_json)

        # Notes
        h3 = H(outlinelevel=2)
        h3.addText('Data Collection Notes')
        doc.text.addElement(h3)

        p_notes = P()
        p_notes.addText(
            'Metrics were collected from the primary production server cluster. '
            'CPU_Usage values represent average utilization percentage. '
            'Memory_Usage values represent total RAM consumed in gigabytes.'
        )
        doc.text.addElement(p_notes)

        doc.save(OUTPUT)
        print(f'ODT created directly with odfpy: {OUTPUT}')
    except Exception as e:
        print(f'odfpy fallback also failed: {e}')
        # Last resort: create minimal ODT structure manually
        create_minimal_odt(json_str)


def create_minimal_odt(json_str):
    """Last resort: create a minimal ODT file as a ZIP archive."""
    import zipfile
    import io

    content_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<office:document-content
  xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
  xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0"
  office:version="1.2">
  <office:body>
    <office:text>
      <text:h text:outline-level="1">Server Performance Data Report</text:h>
      <text:p>This report contains raw server performance metrics collected throughout 2024. The data below is provided in JSON format for programmatic analysis.</text:p>
      <text:h text:outline-level="2">Raw Data (JSON Format)</text:h>
      <text:p>The following JSON array contains 20 statistical records.</text:p>
      <text:p>{json_str.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')}</text:p>
      <text:h text:outline-level="2">Data Collection Notes</text:h>
      <text:p>Metrics were collected from the primary production server cluster. CPU_Usage values represent average utilization percentage. Memory_Usage values represent total RAM consumed in gigabytes.</text:p>
    </office:text>
  </office:body>
</office:document-content>'''

    mimetype = 'application/vnd.oasis.opendocument.text'
    manifest_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0">
  <manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.text"/>
  <manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/>
</manifest:manifest>'''

    with zipfile.ZipFile(OUTPUT, 'w', zipfile.ZIP_DEFLATED) as zf:
        # mimetype must be first and uncompressed
        zf.writestr(zipfile.ZipInfo('mimetype'), mimetype)
        zf.writestr('META-INF/manifest.xml', manifest_xml)
        zf.writestr('content.xml', content_xml)

    print(f'Minimal ODT created manually: {OUTPUT}')


create_initial()
