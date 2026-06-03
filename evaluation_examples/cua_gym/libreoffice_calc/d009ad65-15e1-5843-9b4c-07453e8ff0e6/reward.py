"""
Reward Script: Fetch Docker networking page and save as docker_networking.docx
Task ID: osworld_multi_apps_web_to_doc_008
Domain: multi_apps (Chrome + LibreOffice Writer)
Scoring:
  Component 1 (0.3): docker_networking.docx exists at /home/user/Desktop/ and is non-trivial
  Component 2 (0.3): Document contains 'Networking overview' heading/title (proves correct page)
  Component 3 (0.2): Document contains Docker networking driver names (bridge, host, overlay)
  Component 4 (0.2): Document has substantial content with multiple sections/headings
"""

import os

WORKDIR = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_web_to_doc_008'
FILE_PATH = f'{WORKDIR}/docker_networking.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Component 1: docker_networking.docx exists at /home/user/Desktop/ and is non-trivial (0.3 points)
    # This FAILS on initial_env (empty desktop) and PASSES on golden_env
    try:
        file_size = os.path.getsize(file_path)
        if file_size >= 1000:
            print(f"PASS: Component 1 — docker_networking.docx exists at {file_path} ({file_size} bytes) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — file exists but too small ({file_size} bytes), likely empty/corrupted")
            print(f"\nScore: {total_score}/1.0")
            print(f"REWARD: {total_score}")
            return total_score
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Load docx for remaining components
    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load docx {file_path}: {e}")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Collect all text for searching
    all_text_lower = ' '.join(p.text for p in doc.paragraphs).lower()

    # Component 2: Document contains 'Networking overview' as a heading/title (0.3 points)
    # This is the key content anchor — proves the correct Docker networking page was fetched
    try:
        networking_overview_paras = [
            p for p in doc.paragraphs
            if 'networking overview' in p.text.lower() and p.text.strip()
        ]
        if len(networking_overview_paras) >= 1:
            print(f"PASS: Component 2 — 'Networking overview' heading found in document (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — 'Networking overview' heading not found in document")
            sample = doc.paragraphs[0].text[:80] if doc.paragraphs else 'empty'
            print(f"  (document has {len(doc.paragraphs)} paragraphs, first: {sample!r})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Document contains Docker networking driver names (0.2 points)
    # The Docker networking page lists bridge, host, overlay drivers prominently
    # Presence of all three confirms the right content was fetched
    try:
        required_drivers = ['bridge', 'overlay', 'host']
        found_drivers = [d for d in required_drivers if d in all_text_lower]
        if len(found_drivers) >= 3:
            print(f"PASS: Component 3 — All required Docker network driver names found: {found_drivers} (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — Only {len(found_drivers)}/3 required drivers found: {found_drivers}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Document has substantial content with multiple sections/headings (0.2 points)
    # The Docker networking page has sections: Network drivers, Container networks,
    # Published ports, IP address and hostname, DNS services, Proxy server
    try:
        non_empty_paras = [p for p in doc.paragraphs if p.text.strip()]
        heading_paras = [
            p for p in doc.paragraphs
            if p.style.name.startswith('Heading') and p.text.strip()
        ]
        if len(non_empty_paras) >= 15 and len(heading_paras) >= 3:
            print(f"PASS: Component 4 — Substantial content: {len(non_empty_paras)} non-empty paragraphs,"
                  f" {len(heading_paras)} headings (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 4 — Insufficient content: {len(non_empty_paras)} non-empty paragraphs,"
                  f" {len(heading_paras)} headings (need >=15 paragraphs and >=3 headings)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Run verification — file must exist for any score > 0
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
