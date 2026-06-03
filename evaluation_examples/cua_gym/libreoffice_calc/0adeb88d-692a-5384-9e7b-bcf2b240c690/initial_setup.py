"""
Initial Setup: Create a 4-page Word document memo for PDF conversion task
Task ID: pdf_gf2_025
Domain: pdf (libreoffice_calc listed but actual task is PDF manipulation)
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'pdf_gf2_025'
DOCS_DIR = f'{WORKDIR}/Documents'
OUTPUT = f'{DOCS_DIR}/memo.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Install python-docx on the VM
    subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)

    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    os.makedirs(DOCS_DIR, exist_ok=True)

    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ---- Page 1: Title and Introduction ----
    title = doc.add_heading('Internal Memorandum', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')  # spacer

    # Memo header block
    meta_lines = [
        ('TO:', 'All Department Heads'),
        ('FROM:', 'Elena Vasquez, Chief Operations Officer'),
        ('DATE:', 'March 28, 2025'),
        ('RE:', 'Q1 2025 Operational Review and Q2 Strategic Priorities'),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        run_label = p.add_run(label + ' ')
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_value = p.add_run(value)
        run_value.font.size = Pt(11)

    # Horizontal line effect
    doc.add_paragraph('_' * 72)
    doc.add_paragraph('')

    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This memorandum provides a comprehensive review of our operational '
        'performance during the first quarter of 2025. Overall, the company '
        'achieved a 12.3% increase in revenue compared to Q4 2024, driven '
        'primarily by strong performance in the Enterprise Solutions division '
        'and the successful launch of the Aurora product line in February.'
    )
    doc.add_paragraph(
        'Despite supply chain disruptions in January that affected our '
        'manufacturing output by approximately 8%, the operations team '
        'successfully implemented contingency protocols that limited the '
        'impact on customer deliveries to under 3% of total orders. Our '
        'customer satisfaction scores remained at 94.2%, well above the '
        'industry average of 87.5%.'
    )
    doc.add_paragraph(
        'Key areas requiring attention in Q2 include the integration of '
        'the recently acquired DataStream Analytics platform, the rollout '
        'of our new employee wellness program, and preparation for the '
        'annual compliance audit scheduled for June 2025.'
    )

    # ---- Page 2: Financial Overview ----
    doc.add_page_break()
    doc.add_heading('2. Financial Performance Overview', level=1)
    doc.add_paragraph(
        'The following table summarizes our financial performance across '
        'all divisions during Q1 2025:'
    )

    # Financial data table
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Light List Accent 1'
    headers = ['Division', 'Q1 Revenue', 'Q1 Expenses', 'Net Margin']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    fin_data = [
        ['Enterprise Solutions', '$4,230,000', '$2,845,000', '32.7%'],
        ['Consumer Products', '$2,890,000', '$2,156,000', '25.4%'],
        ['Cloud Services', '$3,145,000', '$1,987,000', '36.8%'],
        ['Professional Services', '$1,670,000', '$1,234,000', '26.1%'],
        ['Research & Development', '$890,000', '$1,450,000', '-62.9%'],
        ['Total', '$12,825,000', '$9,672,000', '24.6%'],
    ]
    for r, row_data in enumerate(fin_data, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    doc.add_paragraph('')
    doc.add_paragraph(
        'Notable highlights include Cloud Services exceeding its quarterly '
        'target by 18%, largely attributable to the migration of three '
        'Fortune 500 clients to our managed infrastructure platform. The '
        'Enterprise Solutions division closed two major contracts with '
        'government agencies totaling $1.2 million in annual recurring revenue.'
    )
    doc.add_paragraph(
        'Research & Development expenses remain elevated due to ongoing '
        'investment in the next-generation Horizon platform, which is '
        'expected to enter beta testing in Q3 2025. The board has approved '
        'an additional $500,000 allocation for accelerated development of '
        'the machine learning capabilities within the platform.'
    )

    # ---- Page 3: Operational Updates ----
    doc.add_page_break()
    doc.add_heading('3. Operational Updates', level=1)

    doc.add_heading('3.1 Supply Chain Improvements', level=2)
    doc.add_paragraph(
        'Following the January disruptions, the logistics team implemented '
        'several critical improvements to our supply chain resilience:'
    )
    items_sc = [
        'Established secondary supplier agreements with three additional '
        'vendors in the Asia-Pacific region, reducing single-source dependency '
        'from 45% to 22%.',
        'Deployed the new inventory management system (IMS-360) across all '
        'four distribution centers, achieving real-time visibility into stock '
        'levels and automated reorder triggers.',
        'Negotiated expedited shipping contracts with FedEx and DHL, reducing '
        'average delivery times for priority orders from 5.2 days to 3.1 days.',
        'Completed risk assessment for all Tier 1 suppliers, identifying and '
        'mitigating 14 high-risk dependencies.',
    ]
    for item in items_sc:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2 Technology Infrastructure', level=2)
    doc.add_paragraph(
        'The IT department completed the migration of our primary data center '
        'to the new co-location facility in Ashburn, Virginia. This migration '
        'improves our disaster recovery capabilities and reduces latency for '
        'East Coast clients by an average of 23 milliseconds. The total cost '
        'of the migration was $340,000, approximately $60,000 under budget.'
    )
    doc.add_paragraph(
        'Cybersecurity enhancements deployed in Q1 include the implementation '
        'of zero-trust network architecture across all corporate systems, '
        'deployment of advanced endpoint detection and response (EDR) tools '
        'on 2,400 employee devices, and completion of the SOC 2 Type II '
        'audit with zero material findings.'
    )

    doc.add_heading('3.3 Human Resources', level=2)
    doc.add_paragraph(
        'Total headcount as of March 31, 2025: 847 employees (up from 812 '
        'at the start of Q1). New hires were concentrated in Engineering '
        '(18 positions), Sales (9 positions), and Customer Success (8 positions). '
        'Employee turnover rate for Q1 was 4.1%, compared to 6.8% industry '
        'average. The new mentorship program launched in February has already '
        'paired 45 junior employees with senior leaders across the organization.'
    )

    # ---- Page 4: Q2 Priorities and Closing ----
    doc.add_page_break()
    doc.add_heading('4. Q2 2025 Strategic Priorities', level=1)
    doc.add_paragraph(
        'Based on Q1 results and evolving market conditions, the following '
        'priorities have been established for Q2 2025:'
    )

    priorities = [
        ('DataStream Integration', 'Complete Phase 2 of the DataStream '
         'Analytics integration by May 15, including data pipeline '
         'consolidation and unified reporting dashboard deployment.'),
        ('Aurora Product Expansion', 'Launch Aurora Professional tier '
         'targeting mid-market enterprises with 50-500 employees. Target: '
         '200 new subscriptions by end of Q2.'),
        ('Compliance Preparation', 'Finalize all documentation and '
         'internal audit procedures for the annual compliance review. '
         'Department heads must submit compliance checklists by May 1.'),
        ('Employee Wellness Program', 'Roll out the comprehensive wellness '
         'initiative including mental health resources, fitness subsidies, '
         'and flexible work arrangement policies across all offices.'),
        ('Customer Success Scaling', 'Implement the new tiered support '
         'model and hire 5 additional Customer Success Managers to '
         'maintain service levels as the client base grows.'),
    ]
    for title, desc in priorities:
        p = doc.add_paragraph()
        run_title = p.add_run(title + ': ')
        run_title.bold = True
        p.add_run(desc)

    doc.add_paragraph('')
    doc.add_heading('5. Closing Remarks', level=1)
    doc.add_paragraph(
        'Q1 2025 has positioned us well for continued growth. The operational '
        'improvements and strategic investments made this quarter provide a '
        'strong foundation for achieving our annual targets. I encourage all '
        'department heads to review the detailed divisional reports and '
        'schedule one-on-one sessions with my office to discuss Q2 execution '
        'plans.'
    )
    # Signature block
    doc.add_paragraph('')
    p_sig = doc.add_paragraph()
    run_sig = p_sig.add_run('Elena Vasquez')
    run_sig.bold = True
    p_sig.add_run(', Chief Operations Officer, Meridian Technologies, Inc.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Ensure no PDF files exist
    for f in ['memo.pdf', 'memo_final.pdf']:
        path = os.path.join(DOCS_DIR, f)
        if os.path.exists(path):
            os.remove(path)
            print(f'Removed pre-existing: {path}')

    # GUI-ready: open the docx in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
