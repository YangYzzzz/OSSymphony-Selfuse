"""
Initial Setup: HashMap Tutorial Task
Task ID: osworld_multi_apps_misc_040
Domain: multi_apps (LibreOffice Writer + Python)

Creates:
  - /home/user/Desktop/HashMap_Tutorial.docx  (tutorial open in LibreOffice Writer)
  - /home/user/Desktop/hashmap.py             (skeleton with stub methods + tests)
"""

import os
import shlex
import subprocess
import time
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

WORKDIR = '/home/user/Desktop'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_tutorial_doc():
    """Create HashMap_Tutorial.docx with detailed open addressing hash map tutorial."""
    doc = Document()

    # Title
    title = doc.add_heading('HashMap Tutorial: Open Addressing Collision Resolution', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_heading('1. Introduction to Hash Maps', level=2)
    doc.add_paragraph(
        'A HashMap (also called a hash table) is a data structure that stores key-value pairs. '
        'It uses a hash function to compute an index into an array of buckets, from which the '
        'desired value can be found. Hash maps offer O(1) average-case time complexity for '
        'insertions, lookups, and deletions.'
    )

    # Open Addressing
    doc.add_heading('2. Open Addressing Collision Resolution', level=2)
    doc.add_paragraph(
        'When two keys hash to the same index (a collision), we need a strategy to resolve it. '
        'Open addressing resolves collisions by probing for the next available slot within the '
        'hash table itself, rather than using separate chaining (linked lists).'
    )

    doc.add_heading('2.1 Linear Probing', level=3)
    doc.add_paragraph(
        'Linear probing is the simplest form of open addressing. When a collision occurs at '
        'index i, we probe indices i+1, i+2, i+3, ... (wrapping around) until an empty slot '
        'is found.'
    )
    p = doc.add_paragraph()
    p.add_run('Probe sequence: ').bold = True
    p.add_run('h(k, i) = (hash(k) + i) % capacity   for i = 0, 1, 2, ...')

    # HashMap Class Design
    doc.add_heading('3. HashMap Class Design', level=2)
    doc.add_paragraph(
        'Our HashMap class uses a fixed-capacity array with linear probing. Each slot can be '
        'in one of three states:'
    )
    doc.add_paragraph('• EMPTY: never been used (None)', style='List Bullet')
    doc.add_paragraph('• OCCUPIED: contains a (key, value) pair', style='List Bullet')
    doc.add_paragraph('• DELETED: previously occupied, now deleted (tombstone sentinel)', style='List Bullet')

    doc.add_paragraph(
        'The DELETED sentinel is critical for correctness. Without it, probing would stop '
        'prematurely at deleted slots, causing get() to incorrectly report keys as missing.'
    )

    # put() method
    doc.add_heading('4. Implementing the put() Method', level=2)
    doc.add_paragraph(
        'The put(key, value) method inserts or updates a key-value pair. Steps:'
    )
    doc.add_paragraph('1. Compute the hash index: idx = hash(key) % capacity', style='List Number')
    doc.add_paragraph('2. Probe linearly until an EMPTY slot or a slot with the same key is found', style='List Number')
    doc.add_paragraph('3. If the same key is found, update its value (overwrite)', style='List Number')
    doc.add_paragraph('4. If an EMPTY slot is found, insert the new (key, value) pair', style='List Number')
    doc.add_paragraph('5. Track the first DELETED slot encountered during probing — you may insert there', style='List Number')
    doc.add_paragraph(
        'Note: When probing, if you encounter a DELETED slot before finding the key, '
        'remember it as a candidate insertion position. If the key is not found, insert '
        'at the first DELETED slot (or EMPTY slot if no DELETED was encountered).'
    )

    p = doc.add_paragraph()
    p.add_run('Example pseudocode for put():').bold = True
    code_put = (
        'def put(self, key, value):\n'
        '    idx = hash(key) % self.capacity\n'
        '    first_deleted = None\n'
        '    for i in range(self.capacity):\n'
        '        probe = (idx + i) % self.capacity\n'
        '        if self.table[probe] is None:  # EMPTY slot\n'
        '            insert_at = first_deleted if first_deleted is not None else probe\n'
        '            self.table[insert_at] = (key, value)\n'
        '            self.size += 1\n'
        '            return\n'
        '        elif self.table[probe] is self.DELETED:  # DELETED tombstone\n'
        '            if first_deleted is None:\n'
        '                first_deleted = probe\n'
        '        elif self.table[probe][0] == key:  # Found existing key\n'
        '            self.table[probe] = (key, value)\n'
        '            return\n'
        '    # Table is full (all occupied or deleted) — insert at first_deleted\n'
        '    if first_deleted is not None:\n'
        '        self.table[first_deleted] = (key, value)\n'
        '        self.size += 1\n'
    )
    p2 = doc.add_paragraph(code_put)
    p2.runs[0].font.name = 'Courier New'
    p2.runs[0].font.size = Pt(9)

    # get() method
    doc.add_heading('5. Implementing the get() Method', level=2)
    doc.add_paragraph(
        'The get(key) method retrieves the value for a given key. Steps:'
    )
    doc.add_paragraph('1. Compute the hash index: idx = hash(key) % capacity', style='List Number')
    doc.add_paragraph('2. Probe linearly until the key is found, an EMPTY slot is reached, or all slots are checked', style='List Number')
    doc.add_paragraph('3. Skip over DELETED (tombstone) slots — do NOT stop at them', style='List Number')
    doc.add_paragraph('4. If the key is found, return its value', style='List Number')
    doc.add_paragraph('5. If an EMPTY slot is reached, raise KeyError (key was never inserted)', style='List Number')

    p = doc.add_paragraph()
    p.add_run('Example pseudocode for get():').bold = True
    code_get = (
        'def get(self, key):\n'
        '    idx = hash(key) % self.capacity\n'
        '    for i in range(self.capacity):\n'
        '        probe = (idx + i) % self.capacity\n'
        '        if self.table[probe] is None:  # EMPTY — key not present\n'
        '            raise KeyError(key)\n'
        '        elif self.table[probe] is self.DELETED:  # Skip tombstone\n'
        '            continue\n'
        '        elif self.table[probe][0] == key:  # Found it\n'
        '            return self.table[probe][1]\n'
        '    raise KeyError(key)\n'
    )
    p2 = doc.add_paragraph(code_get)
    p2.runs[0].font.name = 'Courier New'
    p2.runs[0].font.size = Pt(9)

    # remove() method
    doc.add_heading('6. Implementing the remove() Method', level=2)
    doc.add_paragraph(
        'The remove(key) method deletes a key-value pair. Steps:'
    )
    doc.add_paragraph('1. Compute the hash index: idx = hash(key) % capacity', style='List Number')
    doc.add_paragraph('2. Probe linearly until the key is found or an EMPTY slot is reached', style='List Number')
    doc.add_paragraph('3. Skip DELETED (tombstone) slots', style='List Number')
    doc.add_paragraph('4. If the key is found, replace the slot with the DELETED sentinel and decrement size', style='List Number')
    doc.add_paragraph('5. If an EMPTY slot is reached, raise KeyError (key not found)', style='List Number')

    p = doc.add_paragraph()
    p.add_run('Example pseudocode for remove():').bold = True
    code_remove = (
        'def remove(self, key):\n'
        '    idx = hash(key) % self.capacity\n'
        '    for i in range(self.capacity):\n'
        '        probe = (idx + i) % self.capacity\n'
        '        if self.table[probe] is None:  # EMPTY — key not present\n'
        '            raise KeyError(key)\n'
        '        elif self.table[probe] is self.DELETED:  # Skip tombstone\n'
        '            continue\n'
        '        elif self.table[probe][0] == key:  # Found it\n'
        '            self.table[probe] = self.DELETED\n'
        '            self.size -= 1\n'
        '            return\n'
        '    raise KeyError(key)\n'
    )
    p2 = doc.add_paragraph(code_remove)
    p2.runs[0].font.name = 'Courier New'
    p2.runs[0].font.size = Pt(9)

    # Summary
    doc.add_heading('7. Summary', level=2)
    doc.add_paragraph(
        'Open addressing with linear probing provides an efficient in-place collision resolution '
        'strategy. Key points to remember:'
    )
    doc.add_paragraph('• Always use a DELETED sentinel (tombstone) when removing elements', style='List Bullet')
    doc.add_paragraph('• The get() and remove() methods must skip tombstones but stop at EMPTY slots', style='List Bullet')
    doc.add_paragraph('• The put() method can reuse DELETED slots for new insertions', style='List Bullet')
    doc.add_paragraph('• Load factor should be kept below 0.7 for good performance', style='List Bullet')

    os.makedirs(WORKDIR, exist_ok=True)
    doc_path = f'{WORKDIR}/HashMap_Tutorial.docx'
    doc.save(doc_path)
    print(f'Tutorial created: {doc_path}')
    return doc_path


def create_hashmap_skeleton():
    """Create hashmap.py with stub methods and test code."""
    skeleton_code = '''\
"""
HashMap Implementation using Open Addressing (Linear Probing)
Read HashMap_Tutorial.docx on the Desktop for detailed instructions.
"""


class HashMap:
    """
    A hash map implemented with open addressing (linear probing) collision resolution.

    Internal states for each slot:
      - None  : EMPTY (never used)
      - DELETED sentinel : slot was previously occupied but has been removed
      - (key, value) tuple : OCCUPIED
    """

    class _Deleted:
        """Sentinel class to mark deleted slots (tombstone)."""
        def __repr__(self):
            return '<DELETED>'

    DELETED = _Deleted()

    def __init__(self, capacity=16):
        self.capacity = capacity
        self.table = [None] * self.capacity
        self.size = 0

    def _hash(self, key):
        """Compute the starting probe index for a key."""
        return hash(key) % self.capacity

    def put(self, key, value):
        """
        Insert or update the key-value pair in the hash map.
        Uses linear probing. Can reuse DELETED slots.
        """
        # TODO: Implement this method following the tutorial
        pass

    def get(self, key):
        """
        Return the value associated with key.
        Raises KeyError if the key is not found.
        """
        # TODO: Implement this method following the tutorial
        pass

    def remove(self, key):
        """
        Remove the key-value pair associated with key.
        Raises KeyError if the key is not found.
        """
        # TODO: Implement this method following the tutorial
        pass

    def __len__(self):
        return self.size

    def __contains__(self, key):
        try:
            self.get(key)
            return True
        except KeyError:
            return False


# ============================================================
# Tests
# ============================================================

def run_tests():
    results = []

    # Test 1: Basic put and get
    hm = HashMap()
    hm.put("apple", 1)
    hm.put("banana", 2)
    hm.put("cherry", 3)
    assert hm.get("apple") == 1, "Test 1a failed"
    assert hm.get("banana") == 2, "Test 1b failed"
    assert hm.get("cherry") == 3, "Test 1c failed"
    results.append("Test 1 PASSED: Basic put and get")

    # Test 2: Update existing key
    hm.put("apple", 10)
    assert hm.get("apple") == 10, "Test 2 failed"
    assert len(hm) == 3, "Test 2 size failed"
    results.append("Test 2 PASSED: Update existing key")

    # Test 3: Remove key
    hm.remove("banana")
    assert "banana" not in hm, "Test 3a failed"
    assert len(hm) == 2, "Test 3b size failed"
    results.append("Test 3 PASSED: Remove key")

    # Test 4: KeyError on missing key
    try:
        hm.get("banana")
        results.append("Test 4 FAILED: Expected KeyError for removed key")
    except KeyError:
        results.append("Test 4 PASSED: KeyError on missing key")

    # Test 5: KeyError on remove of missing key
    try:
        hm.remove("durian")
        results.append("Test 5 FAILED: Expected KeyError for absent key")
    except KeyError:
        results.append("Test 5 PASSED: KeyError on remove of absent key")

    # Test 6: Re-insert after remove (tombstone reuse)
    hm.put("banana", 99)
    assert hm.get("banana") == 99, "Test 6 failed"
    assert len(hm) == 3, "Test 6 size failed"
    results.append("Test 6 PASSED: Re-insert after remove (tombstone reuse)")

    # Test 7: Many insertions (stress test)
    hm2 = HashMap(capacity=32)
    for i in range(20):
        hm2.put(f"key{i}", i * 10)
    for i in range(20):
        assert hm2.get(f"key{i}") == i * 10, f"Test 7 failed at key{i}"
    results.append("Test 7 PASSED: Many insertions stress test")

    # Test 8: Contains check
    assert "cherry" in hm, "Test 8a failed"
    assert "nothere" not in hm, "Test 8b failed"
    results.append("Test 8 PASSED: Contains check")

    return results


if __name__ == "__main__":
    import sys
    output_path = "/home/user/Desktop/hashmap_result.txt"
    test_results = run_tests()
    output_lines = []
    output_lines.append("HashMap Test Results")
    output_lines.append("=" * 40)
    for line in test_results:
        output_lines.append(line)
    output_lines.append("=" * 40)
    output_lines.append(f"All {len(test_results)} tests completed.")
    output_text = "\\n".join(output_lines)
    print(output_text)
    with open(output_path, "w") as f:
        f.write(output_text + "\\n")
    print(f"Results written to {output_path}")
'''

    skeleton_path = f'{WORKDIR}/hashmap.py'
    os.makedirs(WORKDIR, exist_ok=True)
    with open(skeleton_path, 'w') as f:
        f.write(skeleton_code)
    print(f'Skeleton created: {skeleton_path}')
    return skeleton_path


def main():
    os.makedirs(WORKDIR, exist_ok=True)

    # Create tutorial document
    doc_path = create_tutorial_doc()

    # Create hashmap skeleton
    create_hashmap_skeleton()

    # Open HashMap_Tutorial.docx in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{doc_path}"', delay_sec=3.0)

    print('GUI_READY: launched LibreOffice Writer with HashMap_Tutorial.docx (DISPLAY=:0)')


main()
