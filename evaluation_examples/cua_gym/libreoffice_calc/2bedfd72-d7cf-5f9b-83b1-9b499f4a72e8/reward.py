"""
Reward Script: Cross-check PDF invoices against bank statement, highlight discrepancies, write memo
Task ID: osworld_multi_apps_doc_pdf_calc_007
Domain: libreoffice_calc + libreoffice_writer (multi-app)

Expected outcome:
1. bank_statement.ods updated: new Sheet2 'Discrepancies' listing VB-002 and VD-004 with red background
2. discrepancy_memo.odt on Desktop: formal memo format with To/From/Date/Subject headers
   listing the 2 missing invoices with amounts

Scoring rubric:
  Component 1: 'Discrepancies' sheet exists in bank_statement.ods           — 0.20 pts
  Component 2: Discrepancies sheet has correct data (VB-002 and VD-004)     — 0.35 pts
  Component 3: Data rows in Discrepancies sheet have red background fill     — 0.15 pts
  Component 4: discrepancy_memo.odt file exists on Desktop                  — 0.10 pts
  Component 5: Memo contains required memo headers and missing invoice refs  — 0.20 pts
  Total: 1.0
"""

import os
import shutil

WORKDIR = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_doc_pdf_calc_007'

BANK_STATEMENT_PATH = os.path.join(WORKDIR, 'bank_statement.ods')
MEMO_PATH = os.path.join(WORKDIR, 'discrepancy_memo.odt')

# Expected missing invoice references and amounts
EXPECTED_REFS = {'VB-002', 'VD-004'}
EXPECTED_AMOUNTS = {
    'VB-002': 875.5,
    'VD-004': 2100.0,
}


def is_red_color(argb_str):
    """Check if a color string represents red (ARGB form FFFF0000)."""
    if argb_str is None:
        return False
    return argb_str.upper().strip() in ('FFFF0000', 'FF0000')


def verify_task():
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # -- Precondition: bank_statement.ods must exist --
    if not os.path.exists(BANK_STATEMENT_PATH):
        print(f"CRITICAL: bank_statement.ods not found at {BANK_STATEMENT_PATH}")
        print("REWARD: 0.0")
        return 0.0

    # Load the bank_statement.ods (it uses xlsx internals despite the .ods extension)
    try:
        import openpyxl
        tmp_path = '/tmp/bank_statement_reward_check.xlsx'
        shutil.copy(BANK_STATEMENT_PATH, tmp_path)
        wb = openpyxl.load_workbook(tmp_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load bank_statement.ods: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: 'Discrepancies' sheet exists in bank_statement.ods (0.20 pts)
    # Initial state: only 'Transactions' sheet. Golden state: 'Discrepancies' added.
    # -------------------------------------------------------------------------
    try:
        sheet_names = wb.sheetnames
        if 'Discrepancies' in sheet_names:
            print(f"PASS: Component 1 — 'Discrepancies' sheet exists (sheets: {sheet_names}) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Expected 'Discrepancies' sheet, found: {sheet_names}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Discrepancies sheet has correct data (VB-002 and VD-004) (0.35 pts)
    # Checks both references and the amounts are present and correct.
    # -------------------------------------------------------------------------
    found_refs = set()
    found_amounts = {}

    try:
        if 'Discrepancies' in wb.sheetnames:
            ws_disc = wb['Discrepancies']
            for row in ws_disc.iter_rows(min_row=1, max_row=20):
                row_str_vals = [str(c.value).strip() if c.value is not None else '' for c in row]
                for ref in EXPECTED_REFS:
                    if ref in row_str_vals:
                        found_refs.add(ref)
                        # Find numeric amount in same row
                        for c in row:
                            if isinstance(c.value, (int, float)):
                                found_amounts[ref] = float(c.value)
                                break

            missing_found = found_refs & EXPECTED_REFS
            if len(missing_found) == 2:
                # Both references found, verify amounts
                amount_ok = all(
                    abs(found_amounts.get(ref, -1) - EXPECTED_AMOUNTS[ref]) < 0.1
                    for ref in EXPECTED_REFS
                )
                if amount_ok:
                    print(f"PASS: Component 2 — Both VB-002 (${found_amounts.get('VB-002')}) and VD-004 (${found_amounts.get('VD-004')}) in Discrepancies sheet (0.35 pts)")
                    total_score += 0.35
                else:
                    print(f"PARTIAL: Component 2 — Both refs found but amounts mismatch (expected VB-002=$875.50, VD-004=$2100.00, got: {found_amounts}) (0.20 pts)")
                    total_score += 0.20
            elif len(missing_found) == 1:
                print(f"PARTIAL: Component 2 — Only one of VB-002/VD-004 found: {missing_found} (0.17 pts)")
                total_score += 0.17
            else:
                print(f"FAIL: Component 2 — Neither VB-002 nor VD-004 found in Discrepancies sheet")
        else:
            print("SKIP: Component 2 — Discrepancies sheet not found, skipping data check")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Data rows in Discrepancies sheet have red background fill (0.15 pts)
    # Red background = fgColor FFFF0000 on data rows (not header row)
    # -------------------------------------------------------------------------
    try:
        if 'Discrepancies' in wb.sheetnames:
            ws_disc = wb['Discrepancies']
            red_rows = 0
            total_data_rows = 0
            for row in ws_disc.iter_rows(min_row=2, max_row=20):
                row_str_vals = [str(c.value).strip() if c.value is not None else '' for c in row]
                if not any(row_str_vals):
                    continue
                if any(ref in row_str_vals for ref in EXPECTED_REFS):
                    total_data_rows += 1
                    # Check if any cell in this row has a red fill
                    row_red_count = sum(
                        1 for cell in row
                        if is_red_color(
                            (lambda c: c.fill.fgColor.rgb if c.fill and c.fill.fgColor else None)(cell)
                        )
                    )
                    if row_red_count > 0:
                        red_rows += 1

            if total_data_rows > 0 and red_rows == total_data_rows:
                print(f"PASS: Component 3 — All {red_rows} discrepancy rows have red background fill (0.15 pts)")
                total_score += 0.15
            elif red_rows > 0 and total_data_rows > 0:
                partial = round(0.15 * red_rows / total_data_rows, 3)
                print(f"PARTIAL: Component 3 — {red_rows}/{total_data_rows} rows have red fill ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — No red background fill found on discrepancy rows (total_data_rows={total_data_rows})")
        else:
            print("SKIP: Component 3 — Discrepancies sheet not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: discrepancy_memo.odt file exists on Desktop (0.10 pts)
    # -------------------------------------------------------------------------
    try:
        memo_exists = os.path.exists(MEMO_PATH) and os.path.getsize(MEMO_PATH) > 100
        if memo_exists:
            print(f"PASS: Component 4 — discrepancy_memo.odt exists at {MEMO_PATH} (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — discrepancy_memo.odt not found or empty at {MEMO_PATH}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -------------------------------------------------------------------------
    # Component 5: Memo contains required content (0.20 pts)
    # Must have memo headers (TO, FROM, SUBJECT) and mention both missing invoices
    # -------------------------------------------------------------------------
    try:
        memo_file_ok = os.path.exists(MEMO_PATH) and os.path.getsize(MEMO_PATH) > 100
        if memo_file_ok:
            from docx import Document
            memo_doc = Document(MEMO_PATH)

            # Collect all text from the document (paragraphs + tables)
            full_text_parts = []
            for para in memo_doc.paragraphs:
                if para.text.strip():
                    full_text_parts.append(para.text.strip())
            for table in memo_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text_parts.append(cell.text.strip())

            full_text = ' '.join(full_text_parts).upper()

            # Check for memo structure (at least 2 of: TO, FROM, SUBJECT, DATE)
            memo_headers = ['TO:', 'FROM:', 'SUBJECT:', 'DATE:']
            headers_found = sum(1 for h in memo_headers if h in full_text)
            has_memo_format = headers_found >= 2

            # Check for both invoice references
            has_vb002 = 'VB-002' in full_text or 'VB002' in full_text
            has_vd004 = 'VD-004' in full_text or 'VD004' in full_text

            # Check for amounts
            has_amount_vb = '875' in full_text
            has_amount_vd = '2100' in full_text or '2,100' in full_text

            print(f"DEBUG Component 5: headers_found={headers_found}, has_vb002={has_vb002}, has_vd004={has_vd004}, has_amount_vb={has_amount_vb}, has_amount_vd={has_amount_vd}")

            if has_memo_format and has_vb002 and has_vd004 and has_amount_vb and has_amount_vd:
                print(f"PASS: Component 5 — Memo has proper format ({headers_found} headers) and references both missing invoices with amounts (0.20 pts)")
                total_score += 0.20
            elif has_memo_format and has_vb002 and has_vd004:
                print(f"PARTIAL: Component 5 — Memo has format and refs but amounts not confirmed (0.13 pts)")
                total_score += 0.13
            elif has_vb002 and has_vd004:
                print(f"PARTIAL: Component 5 — Memo mentions both invoice refs but memo format incomplete (0.10 pts)")
                total_score += 0.10
            elif has_memo_format and (has_vb002 or has_vd004):
                print(f"PARTIAL: Component 5 — Memo has format but only one invoice ref found (0.08 pts)")
                total_score += 0.08
            elif has_vb002 or has_vd004:
                print(f"PARTIAL: Component 5 — Memo mentions at least one invoice ref (0.05 pts)")
                total_score += 0.05
            else:
                print(f"FAIL: Component 5 — Memo missing invoice references or memo structure")
        else:
            print("SKIP: Component 5 — discrepancy_memo.odt not found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score:.4f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


verify_task()
