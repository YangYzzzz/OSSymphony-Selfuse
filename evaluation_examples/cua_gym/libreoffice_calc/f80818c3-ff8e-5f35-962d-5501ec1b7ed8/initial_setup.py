"""
Initial Setup: Quick Sort Reference and Stub Implementation
Task ID: osworld_multi_apps_misc_037
Domain: multi_apps (LibreOffice Writer + Python scripting)

Creates:
  - /home/user/Desktop/Quick_Sort_Reference.docx  (quicksort reference doc)
  - /home/user/Desktop/quicksort.py               (stub with empty implementation)
Opens Quick_Sort_Reference.docx in LibreOffice Writer.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_misc_037'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_reference_doc():
    """Create Quick_Sort_Reference.docx with detailed quicksort explanation and Lomuto pseudocode."""
    doc_path = f'{WORKDIR}/Quick_Sort_Reference.docx'
    doc = Document()

    # Title
    title = doc.add_heading('Quick Sort Reference Guide', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction
    doc.add_heading('Overview', level=1)
    intro = doc.add_paragraph(
        'Quicksort is a highly efficient, comparison-based sorting algorithm. '
        'It uses a divide-and-conquer strategy to sort elements. On average, it '
        'makes O(n log n) comparisons to sort n items. In the worst case, it makes '
        'O(n²) comparisons, but this behavior is rare in practice.'
    )

    doc.add_heading('Lomuto Partition Scheme', level=1)
    doc.add_paragraph(
        'The Lomuto partition scheme selects the last element as the pivot. '
        'It then rearranges elements so that all elements less than the pivot '
        'are to its left, and all elements greater than the pivot are to its right. '
        'The pivot is then placed at its final sorted position.'
    )

    # Pseudocode section
    doc.add_heading('Pseudocode', level=2)
    doc.add_paragraph('Partition function (Lomuto scheme):')

    pseudocode_partition = (
        'function partition(arr, low, high):\n'
        '    pivot = arr[high]        // choose last element as pivot\n'
        '    i = low - 1             // index of smaller element\n'
        '    for j = low to high - 1:\n'
        '        if arr[j] <= pivot:\n'
        '            i = i + 1\n'
        '            swap arr[i] and arr[j]\n'
        '    swap arr[i + 1] and arr[high]\n'
        '    return i + 1            // return pivot\'s final position'
    )

    p = doc.add_paragraph()
    run = p.add_run(pseudocode_partition)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_paragraph('QuickSort recursive function:')

    pseudocode_quicksort = (
        'function quickSort(arr, low, high):\n'
        '    if low < high:\n'
        '        pi = partition(arr, low, high)\n'
        '        quickSort(arr, low, pi - 1)   // recursively sort left sub-array\n'
        '        quickSort(arr, pi + 1, high)  // recursively sort right sub-array'
    )

    p2 = doc.add_paragraph()
    run2 = p2.add_run(pseudocode_quicksort)
    run2.font.name = 'Courier New'
    run2.font.size = Pt(10)

    # Key steps explanation
    doc.add_heading('Algorithm Steps', level=2)
    doc.add_paragraph(
        '1. Choose the last element of the current sub-array as the pivot.',
        style='List Number'
    )
    doc.add_paragraph(
        '2. Partition the array: move all elements smaller than or equal to the pivot to the left, '
        'and all elements greater than the pivot to the right.',
        style='List Number'
    )
    doc.add_paragraph(
        '3. Place the pivot at its correct sorted position (returned by partition).',
        style='List Number'
    )
    doc.add_paragraph(
        '4. Recursively apply quickSort to the left sub-array (low to pi-1).',
        style='List Number'
    )
    doc.add_paragraph(
        '5. Recursively apply quickSort to the right sub-array (pi+1 to high).',
        style='List Number'
    )
    doc.add_paragraph(
        '6. Base case: if low >= high, the sub-array has 0 or 1 elements and is already sorted.',
        style='List Number'
    )

    # Implementation notes
    doc.add_heading('Implementation Notes', level=1)
    doc.add_paragraph(
        'When implementing in Python, the swap operation can be written as:\n'
        '    arr[i], arr[j] = arr[j], arr[i]\n\n'
        'The partition function should return the final position of the pivot element. '
        'The quickSort function should call partition to get the pivot index (pi), '
        'then recursively sort arr[low..pi-1] and arr[pi+1..high].'
    )

    # Example
    doc.add_heading('Example', level=1)
    doc.add_paragraph(
        'Input:  [10, 7, 8, 9, 1, 5]\n'
        'Step 1: pivot = 5, after partition: [1, 5, 8, 9, 7, 10], pi = 1\n'
        'Step 2: Sort left [1] (base case), Sort right [8, 9, 7, 10]\n'
        'Final:  [1, 5, 7, 8, 9, 10]'
    )

    doc.save(doc_path)
    print(f'Reference document created: {doc_path}')
    return doc_path


def create_quicksort_stub():
    """Create quicksort.py with an empty quickSort function and partition stub."""
    py_path = f'{WORKDIR}/quicksort.py'

    stub_code = '''def partition(arr, low, high):
    """
    Partition function using Lomuto partition scheme.
    Selects arr[high] as pivot and places it in its correct position.
    Returns the index of the pivot after partitioning.
    """
    # TODO: Implement Lomuto partition scheme
    pass


def quickSort(arr, low, high):
    """
    Recursive quicksort function.
    Sorts arr[low..high] in-place using the Lomuto partition scheme.
    """
    # TODO: Implement quickSort
    pass


# Test cases
if __name__ == "__main__":
    test1 = [10, 7, 8, 9, 1, 5]
    quickSort(test1, 0, len(test1) - 1)
    print("Sorted array 1:", test1)

    test2 = [64, 34, 25, 12, 22, 11, 90]
    quickSort(test2, 0, len(test2) - 1)
    print("Sorted array 2:", test2)

    test3 = [3, 6, 8, 10, 1, 2, 1]
    quickSort(test3, 0, len(test3) - 1)
    print("Sorted array 3:", test3)

    test4 = [1]
    quickSort(test4, 0, len(test4) - 1)
    print("Sorted array 4:", test4)

    test5 = [5, 4, 3, 2, 1]
    quickSort(test5, 0, len(test5) - 1)
    print("Sorted array 5:", test5)
'''

    with open(py_path, 'w') as f:
        f.write(stub_code)

    print(f'quicksort.py stub created: {py_path}')
    return py_path


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc_path = create_reference_doc()
    create_quicksort_stub()

    # GUI-ready startup: open the reference document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{doc_path}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with Quick_Sort_Reference.docx (DISPLAY=:0)')


create_initial()
