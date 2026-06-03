"""
Reward Script: Board games longest play time lookup task
Task ID: osworld_multi_apps_book_reading_rate_008
Domain: multi_apps (libreoffice_calc + libreoffice_writer)
Scoring:
  Component 1 (0.6): longest_game.docx contains "Twilight Imperium (4th Edition)"
  Component 2 (0.4): boardgames_2023.xlsx has Avg Play Time filled in and Twilight Imperium shows max value
"""

import os

# Domain-specific imports
from docx import Document
import openpyxl

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_book_reading_rate_008'

# Path to the output docx on the Desktop
DOCX_PATH = f'{WORKDIR}/Desktop/longest_game.docx'
# Path to the spreadsheet
XLSX_PATH = f'{WORKDIR}/boardgames_2023.xlsx'

# Expected answer per BGG data
EXPECTED_GAME_NAME = 'Twilight Imperium (4th Edition)'


def verify_task():
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # -----------------------------------------------------------------------
    # Component 1: longest_game.docx contains "Twilight Imperium (4th Edition)"
    # (0.6 points)
    # Initial state: file is empty (0 paragraphs). Golden state: file contains
    # the game name as text. This FAILS on initial and PASSES on golden.
    # -----------------------------------------------------------------------
    try:
        if not os.path.exists(DOCX_PATH):
            print(f"FAIL: Component 1 — {DOCX_PATH} not found")
        else:
            doc = Document(DOCX_PATH)
            # Collect all text from paragraphs and tables
            all_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t:
                            all_text.append(t)
            full_text = ' '.join(all_text)

            # Check: does the document contain the expected game name?
            if EXPECTED_GAME_NAME.lower() in full_text.lower():
                print(f"PASS: Component 1 — longest_game.docx contains '{EXPECTED_GAME_NAME}' (0.6 pts)")
                total_score += 0.6
            else:
                print(f"FAIL: Component 1 — expected '{EXPECTED_GAME_NAME}' in document, found text: {repr(full_text[:200])}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: boardgames_2023.xlsx has Avg Play Time column populated,
    # and Twilight Imperium (4th Edition) has the maximum value.
    # (0.4 points)
    # Initial state: Avg Play Time column is empty (all None).
    # Golden state: all 5 games have numeric values filled in, and TI4 = 240
    # which is the max. This FAILS on initial and PASSES on golden.
    # -----------------------------------------------------------------------
    try:
        if not os.path.exists(XLSX_PATH):
            print(f"FAIL: Component 2 — {XLSX_PATH} not found")
        else:
            wb = openpyxl.load_workbook(XLSX_PATH)
            ws = wb.active

            # Find Avg Play Time column (expected column C = index 3)
            # Build a dict of {game_name: avg_play_time}
            header_row = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]

            # Find column indices
            game_col = None
            time_col = None
            for idx, h in enumerate(header_row, start=1):
                if h and 'game' in str(h).lower():
                    game_col = idx
                if h and ('play time' in str(h).lower() or 'avg' in str(h).lower()):
                    time_col = idx

            if game_col is None or time_col is None:
                print(f"FAIL: Component 2 — Could not find game or play time columns. Headers: {header_row}")
            else:
                game_times = {}
                for row in range(2, ws.max_row + 1):
                    game_name = ws.cell(row=row, column=game_col).value
                    play_time = ws.cell(row=row, column=time_col).value
                    if game_name:
                        game_times[game_name] = play_time

                # Check: all 5 games should have numeric play times
                all_filled = all(
                    isinstance(v, (int, float)) and v > 0
                    for v in game_times.values()
                )
                if not all_filled:
                    print(f"FAIL: Component 2 — Not all Avg Play Time values are filled. Values: {game_times}")
                else:
                    # Check: Twilight Imperium has the maximum value
                    max_time = max(game_times.values())
                    max_games = [g for g, t in game_times.items() if t == max_time]

                    # Twilight Imperium should have max play time (240 min per BGG)
                    ti_present = any(EXPECTED_GAME_NAME.lower() in g.lower() for g in max_games)
                    if ti_present:
                        print(f"PASS: Component 2 — Avg Play Time filled in, max={max_time} min for {max_games} (0.4 pts)")
                        total_score += 0.4
                    else:
                        print(f"FAIL: Component 2 — Expected '{EXPECTED_GAME_NAME}' to have max play time, but max={max_time} for {max_games}")
                        print(f"  All values: {game_times}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


verify_task()
