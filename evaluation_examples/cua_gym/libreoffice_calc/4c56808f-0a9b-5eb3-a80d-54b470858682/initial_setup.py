"""
Initial Setup: Create Reading_list.docx on Desktop and open Chrome with Goodreads
Task ID: osworld_multi_apps_misc_003
Domain: libreoffice_calc (multi-app: Writer + Chrome + Calc)
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_misc_003'
DESKTOP = f'{WORKDIR}/Desktop'
DOCX_OUTPUT = f'{DESKTOP}/Reading_list.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_reading_list():
    """Create the Reading_list.docx file on the Desktop."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Book Club Reading List")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(12)

    # Introduction paragraph
    intro = doc.add_paragraph(
        "Below is the list of books our book club has selected for this season. "
        "Please read the assigned book before our next meeting and come prepared to discuss."
    )
    intro.paragraph_format.space_after = Pt(10)

    # Blank line
    doc.add_paragraph("")

    # Subheading
    heading = doc.add_paragraph()
    heading_run = heading.add_run("Books to Read:")
    heading_run.bold = True
    heading_run.font.size = Pt(13)
    heading.paragraph_format.space_after = Pt(8)

    # The 5 books as a numbered list
    books = [
        "1984",
        "Brave New World",
        "Fahrenheit 451",
        "The Handmaid\u2019s Tale",
        "Never Let Me Go",
    ]

    for book in books:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(book)
        run.font.size = Pt(12)

    # Footer note
    doc.add_paragraph("")
    note = doc.add_paragraph(
        "Note: You can find ratings and reviews for each book on Goodreads (www.goodreads.com)."
    )
    note.paragraph_format.space_before = Pt(10)
    note_run = note.runs[0]
    note_run.font.size = Pt(10)
    note_run.font.italic = True

    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    doc.save(DOCX_OUTPUT)
    print(f'Reading list docx created: {DOCX_OUTPUT}')


def setup_gui():
    """Launch all required apps for the initial GUI state."""
    # 1. Launch Chrome with Goodreads
    launch_gui('google-chrome "https://www.goodreads.com"', delay_sec=3.0)

    # 2. Open Reading_list.docx in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DOCX_OUTPUT}"', delay_sec=2.0)

    print('GUI_READY: launched Chrome (Goodreads) and LibreOffice Writer with DISPLAY=:0')


def main():
    # Install dependencies if needed
    subprocess.run(
        ['pip3', 'install', 'python-docx', '--quiet'],
        capture_output=True
    )

    create_reading_list()
    setup_gui()

    print(f'Initial setup complete.')
    print(f'  - Reading_list.docx: {DOCX_OUTPUT}')
    print(f'  - Chrome open at: https://www.goodreads.com')
    print(f'  - LibreOffice Writer open with Reading_list.docx')


main()
