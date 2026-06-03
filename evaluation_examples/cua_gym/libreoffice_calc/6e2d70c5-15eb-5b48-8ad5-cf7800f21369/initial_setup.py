"""
Initial Setup: Create Must_have_extensions.docx listing six Chrome extensions for remote workers.
Task ID: osworld_multi_apps_misc_011
Domain: multi_apps (LibreOffice Writer + Chrome + OS)
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_misc_011'
DOCX_PATH = f'{WORKDIR}/Must_have_extensions.docx'
DESKTOP_PATH = f'{WORKDIR}/Desktop'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Install python-docx if not available
    subprocess.run(
        ['pip3', 'install', 'python-docx', '--quiet'],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )

    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    # Create Must_have_extensions.docx
    doc = Document()

    # Title
    title = doc.add_heading('Must-Have Chrome Extensions for Remote Workers', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction paragraph
    intro = doc.add_paragraph(
        'The following extensions are essential tools for remote team collaboration, '
        'productivity, and communication. Please ensure all of these are installed in '
        'your Chrome browser before your first day.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # Section heading
    doc.add_heading('Required Extensions', level=2)

    # List of extensions with descriptions
    extensions = [
        ('Loom for Chrome',
         'Record and share video messages instantly. Essential for async communication '
         'across time zones. Replaces long email threads with quick screen recordings.'),
        ('Calendly for Chrome',
         'Schedule meetings without the back-and-forth. Integrates with your calendar '
         'to let teammates and clients book time directly on your schedule.'),
        ('Miro',
         'Collaborative online whiteboard for brainstorming, planning, and visual '
         'thinking. Perfect for remote workshops and sprint planning sessions.'),
        ('Figma',
         'Design and prototype tool with real-time collaboration. Allows developers '
         'and designers to work together on the same files simultaneously.'),
        ('Slack for Chrome',
         'Team messaging and collaboration hub. Keeps all your project communications, '
         'files, and integrations organized in one place for the whole team.'),
        ('Zoom',
         'Video conferencing platform for team meetings, webinars, and one-on-ones. '
         'Required for all scheduled team calls and client presentations.'),
    ]

    for i, (name, description) in enumerate(extensions, 1):
        # Extension number and name as bold
        para = doc.add_paragraph(style='List Number')
        run_name = para.add_run(name)
        run_name.bold = True
        run_name.font.size = Pt(12)

        # Description as normal text in next paragraph
        desc_para = doc.add_paragraph(description)
        desc_para.paragraph_format.left_indent = Pt(36)
        desc_para.paragraph_format.space_after = Pt(8)

    # Footer note
    doc.add_paragraph('')
    note = doc.add_paragraph(
        'Note: Install all extensions from the Chrome Web Store. '
        'After installation, create a text file on the Desktop named '
        '"installed_extensions.txt" listing each extension that was successfully installed.'
    )
    note_run = note.runs[0] if note.runs else note.add_run(note.text)
    note.clear()
    run = note.add_run(
        'Note: Install all extensions from the Chrome Web Store. '
        'After installation, create a text file on the Desktop named '
        '"installed_extensions.txt" listing each extension that was successfully installed.'
    )
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.size = Pt(10)

    doc.save(DOCX_PATH)
    print(f'Initial file created: {DOCX_PATH}')

    # Ensure Desktop directory exists
    os.makedirs(DESKTOP_PATH, exist_ok=True)

    # Make sure installed_extensions.txt does NOT exist (pre-task state)
    ext_txt = os.path.join(DESKTOP_PATH, 'installed_extensions.txt')
    if os.path.exists(ext_txt):
        os.remove(ext_txt)
        print(f'Removed pre-existing: {ext_txt}')

    # GUI-ready startup: open Chrome first, then LibreOffice Writer with the docx
    # Kill any existing Chrome instances to ensure clean state
    subprocess.run(['pkill', '-f', 'chrome'], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    time.sleep(1.5)

    # Launch Chrome with remote debugging enabled
    launch_gui(
        'google-chrome --remote-debugging-port=1337 --no-first-run --no-default-browser-check',
        delay_sec=2.5
    )

    # Launch LibreOffice Writer with the docx
    launch_gui(f'libreoffice --writer "{DOCX_PATH}"', delay_sec=2.0)

    print('GUI_READY: launched Chrome and LibreOffice Writer with DISPLAY=:0')


create_initial()
