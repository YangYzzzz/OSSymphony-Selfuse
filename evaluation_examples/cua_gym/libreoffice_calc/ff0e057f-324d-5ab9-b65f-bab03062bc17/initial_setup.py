"""
Initial Setup: Batch convert .docx files to PDF with logging
Task ID: osworld_multi_apps_doc_batch_convert_008
Domain: libreoffice_calc (multi-app: Writer + Calc + OS)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from odf.opendocument import OpenDocumentSpreadsheet
from odf.table import Table, TableRow, TableCell
from odf.text import P
from odf.style import Style, TextProperties, TableCellProperties

WORKDIR = '/home/user'
DESKTOP = '/home/user/Desktop'
BATCH_DIR = '/home/user/Desktop/batch_docs'
TASK_ID = 'osworld_multi_apps_doc_batch_convert_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.ods'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_docx_file(filepath, title, content_paragraphs):
    """Create a realistic .docx file."""
    doc = Document()
    doc.add_heading(title, 0)
    for para in content_paragraphs:
        doc.add_paragraph(para)
    doc.save(filepath)


def create_initial():
    # 1. Create batch_docs directory
    os.makedirs(BATCH_DIR, exist_ok=True)
    print(f'Created batch_docs directory: {BATCH_DIR}')

    # 2. Create 6 realistic .docx files
    docx_files = [
        {
            'name': 'quarterly_report_q1.docx',
            'title': 'Q1 2025 Quarterly Financial Report',
            'content': [
                'Executive Summary',
                'This report presents the financial results for the first quarter of fiscal year 2025. '
                'Total revenue reached $4.2 million, representing a 12% increase compared to Q1 2024.',
                'Revenue Breakdown',
                'Product Sales: $2,850,000 (67.9% of total revenue)',
                'Service Contracts: $920,000 (21.9% of total revenue)',
                'Consulting: $430,000 (10.2% of total revenue)',
                'Operating Expenses',
                'Personnel: $1,680,000',
                'Infrastructure: $420,000',
                'Marketing: $215,000',
                'Net Income: $1,885,000',
                'Conclusion',
                'The company performed strongly in Q1 2025, exceeding projections by 8%. '
                'We expect continued growth in Q2 driven by new product launches.',
            ]
        },
        {
            'name': 'project_proposal_webapp.docx',
            'title': 'Project Proposal: Customer Portal Web Application',
            'content': [
                'Project Overview',
                'This proposal outlines the development of a new customer-facing web portal to '
                'streamline order management, support requests, and account management.',
                'Objectives',
                '1. Reduce customer support tickets by 40% through self-service capabilities',
                '2. Improve customer satisfaction score from 3.8 to 4.5 (out of 5)',
                '3. Enable real-time order tracking and invoice downloads',
                'Scope of Work',
                'Phase 1 (Months 1-2): Requirements gathering and UI/UX design',
                'Phase 2 (Months 3-5): Backend API development and database design',
                'Phase 3 (Months 6-7): Frontend development and integration testing',
                'Phase 4 (Month 8): UAT, security audit, and production deployment',
                'Budget Estimate',
                'Development: $185,000',
                'Infrastructure setup: $28,000',
                'Testing and QA: $22,000',
                'Total estimated budget: $235,000',
            ]
        },
        {
            'name': 'hr_policy_remote_work.docx',
            'title': 'Remote Work Policy 2025',
            'content': [
                'Policy Statement',
                'This policy establishes guidelines for employees working remotely. '
                'Effective January 1, 2025, all eligible employees may work remotely up to 3 days per week.',
                'Eligibility',
                'Full-time employees who have been with the company for at least 6 months are eligible.',
                'Roles requiring physical presence (warehouse, facilities) are not eligible.',
                'Equipment and Security',
                'Company laptops must be used for all work activities.',
                'VPN connection is mandatory when accessing internal systems.',
                'Home internet must meet minimum 50 Mbps download speed.',
                'Core Hours',
                'All remote employees must be available 10:00 AM - 3:00 PM in their local timezone.',
                'Video camera must be on during team meetings and client calls.',
                'Workspace Requirements',
                'Employees must have a dedicated workspace free from distractions.',
                'Ergonomic setup guidelines must be followed to prevent injury.',
            ]
        },
        {
            'name': 'technical_spec_api_v2.docx',
            'title': 'Technical Specification: REST API Version 2.0',
            'content': [
                'Overview',
                'This document describes the technical specification for REST API v2.0, '
                'replacing the legacy SOAP-based API that will be deprecated on December 31, 2025.',
                'Authentication',
                'All API requests require Bearer token authentication using OAuth 2.0.',
                'Tokens expire after 3600 seconds and must be refreshed using the /auth/refresh endpoint.',
                'Rate Limiting',
                'Standard tier: 1,000 requests per hour',
                'Premium tier: 10,000 requests per hour',
                'Enterprise tier: Unlimited (fair use policy applies)',
                'Endpoints',
                'GET /api/v2/products - List all products with pagination',
                'POST /api/v2/orders - Create new order',
                'GET /api/v2/orders/{id} - Get order details',
                'PUT /api/v2/orders/{id}/cancel - Cancel an order',
                'Response Format',
                'All responses are JSON with standard envelope: {status, data, meta, errors}',
                'HTTP status codes follow RFC 7231 conventions.',
            ]
        },
        {
            'name': 'meeting_minutes_2025_03.docx',
            'title': 'Board Meeting Minutes - March 2025',
            'content': [
                'Meeting Details',
                'Date: March 5, 2025',
                'Location: Conference Room A, HQ Building',
                'Attendees: Sarah Chen (CEO), Marcus Johnson (CFO), Emily Rodriguez (CTO), '
                'David Park (COO), Jennifer Walsh (Board Chair)',
                'Agenda Item 1: Q4 2024 Financial Review',
                'CFO Marcus Johnson presented Q4 results showing revenue of $4.8M against target of $4.5M.',
                'Full year 2024 revenue: $17.2M, up 18% YoY.',
                'Motion to approve Q4 financial statements: Passed unanimously.',
                'Agenda Item 2: 2025 Strategic Initiatives',
                'CTO Emily Rodriguez presented the technology roadmap for 2025.',
                'Key initiatives: AI integration, cloud migration, API v2.0 launch.',
                'Budget approved for technology initiatives: $2.1M.',
                'Agenda Item 3: New Market Expansion',
                'COO David Park presented analysis of APAC market opportunity.',
                'Estimated TAM: $85M. Proposed entering Singapore market in Q3 2025.',
                'Board approved initial investment of $500K for market entry.',
                'Next Meeting: June 4, 2025',
            ]
        },
        {
            'name': 'training_guide_onboarding.docx',
            'title': 'Employee Onboarding Guide 2025',
            'content': [
                'Welcome to the Team!',
                'This guide will help you get started at our company. Please read it carefully '
                'and complete all required steps within your first two weeks.',
                'Week 1: Administrative Setup',
                'Day 1: Complete HR paperwork, collect access badge, set up workstation',
                'Day 2: IT setup - email, Slack, project management tools, VPN access',
                'Day 3-4: Department introduction meetings and system access verification',
                'Day 5: Meet with your manager to review 30/60/90-day goals',
                'Required Training Modules',
                'Security Awareness Training (mandatory, complete by end of Week 1)',
                'Data Privacy and GDPR Compliance (mandatory, 2 hours)',
                'Code of Conduct and Ethics (mandatory, 1 hour)',
                'Department-specific technical training (scheduled with team lead)',
                'Key Systems and Tools',
                'Email: Outlook 365 - contact IT for setup',
                'Project Management: Jira - training video in knowledge base',
                'Documentation: Confluence - ask your buddy for a walkthrough',
                'HR Portal: access at hr.company.internal for leave requests and payslips',
                'Your First 90 Days',
                'Days 1-30: Orientation, training, and shadowing existing team members',
                'Days 31-60: Taking on small independent tasks with mentorship',
                'Days 61-90: Full productivity, working on team projects independently',
            ]
        },
    ]

    for docx_info in docx_files:
        filepath = os.path.join(BATCH_DIR, docx_info['name'])
        create_docx_file(filepath, docx_info['title'], docx_info['content'])
        print(f'Created: {filepath}')

    # 3. Create conversion_log.ods with headers only (no data rows)
    doc = OpenDocumentSpreadsheet()
    table = Table(name="conversion_log")

    # Create header row
    header_row = TableRow()
    headers = ['Filename', 'Original_Size', 'PDF_Size', 'Status', 'Timestamp']
    for header in headers:
        cell = TableCell()
        cell.addElement(P(text=header))
        header_row.addElement(cell)
    table.addElement(header_row)

    doc.spreadsheet.addElement(table)
    doc.save(OUTPUT)
    print(f'Initial ODS file created (headers only): {OUTPUT}')

    # 4. GUI-ready startup: open file manager showing batch_docs and the ODS log file
    # Open Nautilus file manager to show the batch_docs folder
    launch_gui(f'nautilus "{BATCH_DIR}"', delay_sec=2.0)
    # Open LibreOffice Calc with the conversion log
    launch_gui(f'libreoffice --calc "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched Nautilus and LibreOffice Calc with DISPLAY=:0')


create_initial()
