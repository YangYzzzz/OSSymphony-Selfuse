"""
Reward Script: Conference Acceptance Rate Analysis
Task ID: osworld_multi_apps_web_conference_013
Domain: libreoffice_calc + libreoffice_writer (multi-app)
Scoring:
  - Component 1: ODS file exists at Desktop with correct 5-column header row (0.25 pts)
  - Component 2: ODS file has 30 data rows with all 5 conferences x 6 years (0.25 pts)
  - Component 3: ODS file has conditional formatting on Rate column (3 rules) (0.20 pts)
  - Component 4: ODT file exists in Documents with 5x6 acceptance rate table (0.20 pts)
  - Component 5: ODT file has analysis/trends paragraph (0.10 pts)
  Total: 1.00
"""

import os
import zipfile
import xml.etree.ElementTree as ET
import re

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_web_conference_013'

ODS_PATH = os.path.join(WORKDIR, 'Desktop', 'acceptance_rates.ods')
ODT_PATH = os.path.join(WORKDIR, 'Documents', 'acceptance_analysis.odt')

EXPECTED_CONFERENCES = {'NeurIPS', 'ICML', 'ICLR', 'ACL', 'CVPR'}
EXPECTED_YEARS = {2019, 2020, 2021, 2022, 2023, 2024}


def verify_task():
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # ===========================================================
    # Component 1: ODS file exists with correct 5-column headers
    # (0.25 points)
    # FAILS on initial_env (no file), PASSES on golden_env
    # ===========================================================
    try:
        if not os.path.exists(ODS_PATH):
            print(f"FAIL: Component 1 — ODS file not found at {ODS_PATH}")
        else:
            # The .ods file is in OOXML (xlsx) format internally — parse via zipfile
            with zipfile.ZipFile(ODS_PATH, 'r') as z:
                sheet_xml = z.read('xl/worksheets/sheet1.xml').decode('utf-8')

            # Parse XML to extract headers from row 1
            ns = {'x': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
            root = ET.fromstring(sheet_xml)
            rows = root.findall('.//x:row', ns)

            if not rows:
                print("FAIL: Component 1 — ODS file has no rows")
            else:
                header_row = rows[0]
                cells = header_row.findall('x:c', ns)
                header_values = []
                for cell in cells:
                    t_elem = cell.find('.//x:t', ns)
                    v_elem = cell.find('x:v', ns)
                    if t_elem is not None:
                        header_values.append(t_elem.text.strip() if t_elem.text else '')
                    elif v_elem is not None:
                        header_values.append(v_elem.text.strip() if v_elem.text else '')

                expected_headers_lower = ['conference', 'year', 'submitted', 'accepted', 'rate']
                actual_lower = [h.lower() for h in header_values]

                if all(exp in actual_lower for exp in expected_headers_lower):
                    print(f"PASS: Component 1 — ODS has correct 5-column headers: {header_values} (0.25 pts)")
                    total_score += 0.25
                else:
                    print(f"FAIL: Component 1 — Expected headers {expected_headers_lower}, found {header_values}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ===========================================================
    # Component 2: ODS has 30 data rows covering all 5 conferences x 6 years
    # (0.25 points)
    # FAILS on initial_env (no file), PASSES on golden_env
    # ===========================================================
    try:
        if not os.path.exists(ODS_PATH):
            print(f"FAIL: Component 2 — ODS file not found at {ODS_PATH}")
        else:
            with zipfile.ZipFile(ODS_PATH, 'r') as z:
                sheet_xml = z.read('xl/worksheets/sheet1.xml').decode('utf-8')

            ns = {'x': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
            root = ET.fromstring(sheet_xml)
            rows = root.findall('.//x:row', ns)

            # Data rows are rows 2+ (skip header row 1)
            data_rows = [r for r in rows if int(r.attrib.get('r', '0')) >= 2]
            data_row_count = len(data_rows)

            # Extract conference names (column A = first cell) and years (column B)
            found_conferences = set()
            found_years = set()
            for row in data_rows:
                cells = row.findall('x:c', ns)
                conf_val = None
                year_val = None
                for cell in cells:
                    ref = cell.attrib.get('r', '')
                    if ref.startswith('A'):
                        t_elem = cell.find('.//x:t', ns)
                        if t_elem is not None and t_elem.text:
                            conf_val = t_elem.text.strip()
                    elif ref.startswith('B'):
                        v_elem = cell.find('x:v', ns)
                        if v_elem is not None and v_elem.text:
                            try:
                                year_val = int(float(v_elem.text))
                            except ValueError:
                                pass
                if conf_val:
                    found_conferences.add(conf_val)
                if year_val:
                    found_years.add(year_val)

            missing_confs = EXPECTED_CONFERENCES - found_conferences
            missing_years = EXPECTED_YEARS - found_years

            if data_row_count == 30 and not missing_confs and not missing_years:
                print(f"PASS: Component 2 — ODS has 30 data rows covering all 5 conferences x 6 years (0.25 pts)")
                total_score += 0.25
            elif data_row_count >= 25 and not missing_confs:
                # Partial: has all conferences but possibly fewer rows
                print(f"PARTIAL: Component 2 — ODS has {data_row_count} rows (expected 30), all conferences present but missing years: {missing_years}")
                if data_row_count >= 25:  # Award partial credit when close to complete
                    total_score += 0.15
            else:
                print(f"FAIL: Component 2 — ODS has {data_row_count} rows, "
                      f"found conferences: {found_conferences}, missing: {missing_confs}, "
                      f"found years: {found_years}, missing: {missing_years}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ===========================================================
    # Component 3: ODS has conditional formatting on Rate column (E)
    # 3 rules: green for >30%, yellow for 20-30%, red for <20%
    # (0.20 points)
    # FAILS on initial_env (no file), PASSES on golden_env
    # ===========================================================
    try:
        if not os.path.exists(ODS_PATH):
            print(f"FAIL: Component 3 — ODS file not found at {ODS_PATH}")
        else:
            with zipfile.ZipFile(ODS_PATH, 'r') as z:
                sheet_xml = z.read('xl/worksheets/sheet1.xml').decode('utf-8')
                styles_xml = z.read('xl/styles.xml').decode('utf-8')

            # Check conditional formatting exists in sheet
            cf_present = 'conditionalFormatting' in sheet_xml

            # Check that sqref applies to column E (Rate column)
            cf_in_e_col = bool(re.search(r'sqref="E\d+:E\d+"', sheet_xml) or
                               re.search(r'sqref="E\d+"', sheet_xml))

            # Check number of cfRule elements (expect at least 3)
            cf_rule_count = sheet_xml.count('<cfRule')

            # Check styles.xml has dxf fills (differential formats for CF colors)
            dxf_count = styles_xml.count('<dxf>')

            # Verify the 3 expected CF thresholds are present
            has_gt30 = '0.30' in sheet_xml or '0.3<' in sheet_xml or '>0.30' in sheet_xml or \
                       bool(re.search(r'greaterThan.*0\.3', sheet_xml))
            has_lt20 = '0.20' in sheet_xml or bool(re.search(r'lessThan.*0\.2', sheet_xml))

            if cf_present and cf_in_e_col and cf_rule_count >= 3 and dxf_count >= 3:
                print(f"PASS: Component 3 — ODS has {cf_rule_count} CF rules on Rate column "
                      f"with {dxf_count} differential formats (0.20 pts)")
                total_score += 0.20
            elif cf_present and cf_rule_count >= 2:
                print(f"PARTIAL: Component 3 — ODS has CF formatting but only {cf_rule_count} rules "
                      f"(expected 3), dxf count: {dxf_count}")
                total_score += 0.10
            else:
                print(f"FAIL: Component 3 — cf_present={cf_present}, cf_rule_count={cf_rule_count}, "
                      f"dxf_count={dxf_count}, cf_in_e_col={cf_in_e_col}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ===========================================================
    # Component 4: ODT file exists in Documents with 5x6 table
    # (conference rows x year columns)
    # (0.20 points)
    # FAILS on initial_env (no file), PASSES on golden_env
    # ===========================================================
    try:
        if not os.path.exists(ODT_PATH):
            print(f"FAIL: Component 4 — ODT file not found at {ODT_PATH}")
        else:
            # The .odt is actually in docx (OOXML) format internally
            from docx import Document as DocxDocument
            doc = DocxDocument(ODT_PATH)

            tables = doc.tables
            if not tables:
                print("FAIL: Component 4 — ODT has no tables")
            else:
                # Find the acceptance rates table: expect 6 rows (header + 5 conferences) x 7 cols (Conference + 6 years)
                found_table = None
                for table in tables:
                    nrows = len(table.rows)
                    ncols = len(table.columns)
                    # Expected: 6 rows x 7 cols for conference x year table
                    if nrows >= 5 and ncols >= 6:
                        # Check header row has year values
                        header_texts = [cell.text.strip() for cell in table.rows[0].cells]
                        if any(str(y) in header_texts for y in EXPECTED_YEARS):
                            found_table = table
                            break

                if found_table is not None:
                    nrows = len(found_table.rows)
                    ncols = len(found_table.columns)
                    # Verify conference names in first column (data rows)
                    conf_in_table = set()
                    for row in found_table.rows[1:]:
                        conf_cell = row.cells[0].text.strip()
                        if conf_cell in EXPECTED_CONFERENCES:
                            conf_in_table.add(conf_cell)

                    missing_confs_in_table = EXPECTED_CONFERENCES - conf_in_table
                    if not missing_confs_in_table and nrows >= 6 and ncols >= 7:
                        print(f"PASS: Component 4 — ODT has {nrows}x{ncols} table with all 5 conferences (0.20 pts)")
                        total_score += 0.20
                    elif not missing_confs_in_table:
                        print(f"PASS: Component 4 — ODT has {nrows}x{ncols} table with all 5 conferences (0.20 pts)")
                        total_score += 0.20
                    else:
                        print(f"FAIL: Component 4 — Table found ({nrows}x{ncols}) but missing conferences: "
                              f"{missing_confs_in_table}")
                else:
                    # Check any table with percentage values
                    table = tables[0]
                    nrows = len(table.rows)
                    ncols = len(table.columns)
                    # Check if rate data present (percentage symbols)
                    rate_cells = 0
                    for row in table.rows[1:]:
                        for cell in row.cells[1:]:
                            if '%' in cell.text:
                                rate_cells += 1
                    if rate_cells >= 20:
                        print(f"PASS: Component 4 — ODT table has {rate_cells} cells with rates (0.20 pts)")
                        total_score += 0.20
                    else:
                        print(f"FAIL: Component 4 — No proper acceptance rate table found. "
                              f"Best table: {nrows}x{ncols}, rate cells: {rate_cells}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ===========================================================
    # Component 5: ODT has an analysis/trends paragraph
    # (mentions conferences and trends/outliers)
    # (0.10 points)
    # FAILS on initial_env (no file), PASSES on golden_env
    # ===========================================================
    try:
        if not os.path.exists(ODT_PATH):
            print(f"FAIL: Component 5 — ODT file not found at {ODT_PATH}")
        else:
            from docx import Document as DocxDocument
            doc = DocxDocument(ODT_PATH)

            # Collect all paragraph text
            full_text = ' '.join(para.text for para in doc.paragraphs if para.text.strip())
            full_text_lower = full_text.lower()

            # Check for analysis keywords: trends/outliers + at least 2 conference names
            has_analysis_keywords = any(kw in full_text_lower for kw in
                                        ['trend', 'outlier', 'analysis', 'notably', 'significant',
                                         'highest', 'lowest', 'increased', 'decreased'])
            conf_mentions = sum(1 for conf in EXPECTED_CONFERENCES if conf.lower() in full_text_lower)

            # Must have more than just a heading and table — require at least 50 chars of analysis text
            # outside of the table itself
            non_empty_paras = [p for p in doc.paragraphs if len(p.text.strip()) > 30]
            long_para_found = any(len(p.text) > 100 for p in non_empty_paras)

            if has_analysis_keywords and conf_mentions >= 2 and long_para_found:
                print(f"PASS: Component 5 — ODT has analysis paragraph with trends/outliers "
                      f"mentioning {conf_mentions} conferences (0.10 pts)")
                total_score += 0.10
            elif long_para_found and conf_mentions >= 1:
                print(f"PARTIAL: Component 5 — ODT has text with {conf_mentions} conference mentions "
                      f"but limited analysis keywords")
                total_score += 0.05
            else:
                print(f"FAIL: Component 5 — analysis_keywords={has_analysis_keywords}, "
                      f"conf_mentions={conf_mentions}, long_para={long_para_found}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


verify_task()
