"""
Initial Setup: Multi-app task - PDF paper + LibreOffice Writer + Chrome browsing
Task ID: osworld_multi_apps_paper_scholar_browse_014
Domain: multi_apps (libreoffice_writer + chrome + pdf)

Creates:
  1. A PDF of a distributed systems paper on the Desktop
     - Paper: "Chord: A Scalable Peer-to-peer Lookup Service for Internet Applications"
     - Corresponding author: Ion Stoica (istoica@cs.berkeley.edu)
  2. A blank LibreOffice Writer document open
  3. Chrome available for browsing Google Scholar
"""

import os
import shlex
import subprocess
import time

from fpdf import FPDF

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_paper_scholar_browse_014'
DESKTOP = f'{WORKDIR}/Desktop'
PDF_PATH = f'{DESKTOP}/paper.pdf'
WRITER_DOC = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_pdf():
    """Create a realistic distributed systems paper PDF on the Desktop."""
    os.makedirs(DESKTOP, exist_ok=True)

    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.multi_cell(0, 8, "Chord: A Scalable Peer-to-peer Lookup Service\nfor Internet Applications", align="C")
    pdf.ln(4)

    # Authors
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 6,
        "Ion Stoica, Robert Morris, David Karger, M. Frans Kaashoek, Hari Balakrishnan",
        align="C")
    pdf.ln(2)

    # Affiliations
    pdf.set_font("Helvetica", style="I", size=10)
    pdf.multi_cell(0, 5,
        "MIT Laboratory for Computer Science\n"
        "chord@lcs.mit.edu",
        align="C")
    pdf.ln(4)

    # Corresponding author note
    pdf.set_font("Helvetica", size=9)
    pdf.set_text_color(80, 80, 80)
    pdf.multi_cell(0, 5,
        "*Corresponding author: Ion Stoica  <istoica@cs.berkeley.edu>",
        align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(6)

    # Horizontal rule via line
    pdf.set_draw_color(100, 100, 100)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(4)

    # Abstract
    pdf.set_font("Helvetica", style="B", size=11)
    pdf.cell(0, 6, "Abstract", ln=True)
    pdf.set_font("Helvetica", size=10)
    pdf.multi_cell(0, 5,
        "A fundamental problem that confronts peer-to-peer applications is to efficiently locate "
        "the node that stores a particular data item. This paper presents Chord, a distributed "
        "lookup protocol that addresses this problem. Chord provides support for just one "
        "operation: given a key, it maps the key onto a node. Data location can be easily "
        "implemented on top of Chord by associating a key with each data item, and storing the "
        "key/data item pair at the node to which the key maps. Chord adapts efficiently as nodes "
        "join and leave the system, and can answer queries even if the system is continuously "
        "changing. Results from theoretical analysis, simulations, and experiments show that "
        "Chord is scalable, with communication cost and the state maintained by each node "
        "scaling as O(log N), where N is the number of Chord nodes.",
        align="J")
    pdf.ln(5)

    # 1. Introduction
    pdf.set_font("Helvetica", style="B", size=11)
    pdf.cell(0, 6, "1. Introduction", ln=True)
    pdf.set_font("Helvetica", size=10)
    pdf.multi_cell(0, 5,
        "Peer-to-peer systems and applications are distributed systems without any centralized "
        "control or hierarchical organization. A key and difficult problem in building an "
        "efficient peer-to-peer application is to implement a distributed lookup that efficiently "
        "routes messages to the nodes responsible for specific data items.\n\n"
        "Chord provides a peer-to-peer lookup service: given a key, Chord maps the key to a "
        "node. Chord uses consistent hashing to assign keys to nodes. With high probability, the "
        "hash function balances load so that each node is responsible for only O(1/N) of the "
        "keys, where N is the number of nodes.",
        align="J")
    pdf.ln(4)

    # 2. System Model
    pdf.set_font("Helvetica", style="B", size=11)
    pdf.cell(0, 6, "2. System Model", ln=True)
    pdf.set_font("Helvetica", size=10)
    pdf.multi_cell(0, 5,
        "The Chord protocol specifies how to find the locations of keys, how new nodes join the "
        "system, and how to recover from the failure (or planned departure) of existing nodes. "
        "Each Chord node needs routing information about only a few other nodes. Because the "
        "routing table is small, a Chord node can join or leave the network with only a small "
        "number of messages.",
        align="J")
    pdf.ln(4)

    # 3. The Chord Protocol
    pdf.set_font("Helvetica", style="B", size=11)
    pdf.cell(0, 6, "3. The Chord Protocol", ln=True)
    pdf.set_font("Helvetica", size=10)
    pdf.multi_cell(0, 5,
        "Chord assigns both keys and nodes a number in an identifier space with 2^m identifiers. "
        "Chord maps keys onto nodes. Key k is assigned to the first node whose identifier is "
        "equal to or follows k in the identifier space. This node is called the successor of k.\n\n"
        "To speed up the process, each Chord node maintains a routing table with at most m "
        "entries (finger table). The i-th entry in the table at node n contains the identity of "
        "the first node s that succeeds n by at least 2^(i-1) on the identifier circle. This "
        "node is called the i-th finger of node n.",
        align="J")
    pdf.ln(4)

    # 4. Conclusion
    pdf.set_font("Helvetica", style="B", size=11)
    pdf.cell(0, 6, "4. Conclusion", ln=True)
    pdf.set_font("Helvetica", size=10)
    pdf.multi_cell(0, 5,
        "Chord is a scalable and efficient protocol for peer-to-peer lookup. Our results show "
        "that Chord can find the owner of a key in O(log N) hops, requires O(log N) state per "
        "node, and that communication overhead during membership changes is O(log^2 N) messages "
        "per joining or leaving node. We believe these features make Chord an attractive building "
        "block for peer-to-peer applications.",
        align="J")
    pdf.ln(4)

    # References
    pdf.set_font("Helvetica", style="B", size=11)
    pdf.cell(0, 6, "References", ln=True)
    pdf.set_font("Helvetica", size=9)
    pdf.multi_cell(0, 5,
        "[1] D. R. Karger, E. Lehman, T. Leighton, M. Levine, D. Lewin, and R. Panigrahy. "
        "Consistent hashing and random trees: Distributed caching protocols for relieving hot "
        "spots on the World Wide Web. STOC 1997.\n"
        "[2] I. Stoica, R. Morris, D. Liben-Nowell, D. Karger, M. F. Kaashoek, F. Dabek, and "
        "H. Balakrishnan. Chord: A scalable peer-to-peer lookup protocol for internet "
        "applications. IEEE/ACM Transactions on Networking, 11(1):17-32, 2003.\n"
        "[3] A. Rowstron and P. Druschel. Pastry: Scalable, distributed object location and "
        "routing for large-scale peer-to-peer systems. Middleware 2001.\n"
        "[4] B. Zhao, J. Kubiatowicz, and A. Joseph. Tapestry: An infrastructure for fault-"
        "tolerant wide-area location and routing. Technical Report UCB/CSD-01-1141, 2001.",
        align="J")

    pdf.output(PDF_PATH)
    print(f"PDF created: {PDF_PATH}")


def create_writer_document():
    """Create a blank Writer document for the agent to fill in."""
    from docx import Document

    doc = Document()
    # Add a title paragraph as a hint
    title = doc.add_paragraph()
    run = title.add_run("Co-authors of Corresponding Author")
    run.bold = True
    from docx.shared import Pt
    run.font.size = Pt(14)

    # Blank lines for the agent to fill in
    doc.add_paragraph("")
    doc.add_paragraph("1. ")
    doc.add_paragraph("2. ")
    doc.add_paragraph("3. ")

    doc.save(WRITER_DOC)
    print(f"Writer document created: {WRITER_DOC}")


def main():
    # Create PDF paper on Desktop
    create_pdf()

    # Create blank Writer document
    create_writer_document()

    # Open the PDF with evince (default PDF viewer) on the Desktop
    launch_gui(f'evince "{PDF_PATH}"', delay_sec=2.0)

    # Open Writer document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{WRITER_DOC}"', delay_sec=2.0)

    # Chrome is available via google-chrome command; agent will open it themselves
    # or we can open it to the new tab page for convenience
    launch_gui('google-chrome', delay_sec=2.0)

    print("GUI_READY: launched PDF viewer, LibreOffice Writer, and Chrome with DISPLAY=:0")


main()
