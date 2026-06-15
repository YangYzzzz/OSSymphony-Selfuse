"""
Reward Script: Fix APA 7th Edition Reference List Formatting in essay.docx
Task ID: osworld_multi_apps_misc_045
Domain: libreoffice_writer
Scoring:
  - Component 1: APA author name format (Last, F. F.,) — 0.35 pts
  - Component 2: DOI presence (https://doi.org/) — 0.35 pts
  - Component 3: Journal/book title italicization — 0.30 pts
Total: 1.0
"""

import os
import re

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_misc_045'
FILE_PATH = '/home/user/Desktop/student_submissions/essay.docx'


def verify_task(file_path):
    """
    Verify that the APA 7th edition reference list has been corrected.
    Checks: (1) author name format, (2) DOI presence, (3) journal title italics.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file " + str(file_path) + ": " + str(e))
        print("REWARD: 0.0")
        return 0.0

    # Locate the References section
    ref_start = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.lower() == 'references':
            ref_start = i
            break

    if ref_start is None:
        print("CRITICAL: Could not find 'References' section in document")
        print("REWARD: 0.0")
        return 0.0

    # Collect all reference paragraphs (non-empty after 'References' heading)
    ref_paras = []
    for i in range(ref_start + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text:
            ref_paras.append(para)

    if len(ref_paras) == 0:
        print("CRITICAL: No reference entries found after 'References' heading")
        print("REWARD: 0.0")
        return 0.0

    total_refs = len(ref_paras)
    print("INFO: Found " + str(total_refs) + " reference entries")

    # -----------------------------------------------------------------------
    # Component 1: APA Author Name Format — 0.35 points
    # APA 7th edition requires author names in format: Last, F. M.,
    # Pattern: starts with LastName, Initial(s).
    # This FAILS on initial (many refs use first-name-first order)
    # and PASSES on golden (all use last-name-first).
    # -----------------------------------------------------------------------
    try:
        apa_author_pattern = re.compile(r'^[A-Z][A-Za-z\u00C0-\u017E\-]+,\s+[A-Z]\.')
        correct_author_count = 0

        for para in ref_paras:
            text = para.text.strip()
            if apa_author_pattern.match(text):
                correct_author_count += 1

        # Golden: 13/14 pass (Al-Mukhaini uses hyphenated name — regex handles it)
        # Threshold: at least 12 out of 14 in correct format
        author_ratio = correct_author_count / total_refs
        if author_ratio >= 12.0 / 14.0:
            print("PASS: Component 1 — APA author format: " + str(correct_author_count) + "/" + str(total_refs) + " refs correct (0.35 pts)")
            total_score += 0.35
        else:
            print("FAIL: Component 1 — APA author format: only " + str(correct_author_count) + "/" + str(total_refs) + " refs in last-name-first format")
    except Exception as e:
        print("ERROR: Component 1 — " + str(e))

    # -----------------------------------------------------------------------
    # Component 2: DOI Presence — 0.35 points
    # APA 7th edition requires DOIs in format https://doi.org/...
    # Initial has only 6/14 DOIs; golden has 14/14.
    # Threshold: at least 12 out of 14 have DOIs.
    # -----------------------------------------------------------------------
    try:
        doi_count = 0
        for para in ref_paras:
            text = para.text.strip()
            if 'https://doi.org/' in text:
                doi_count += 1

        doi_ratio = doi_count / total_refs
        if doi_ratio >= 12.0 / 14.0:
            print("PASS: Component 2 — DOI presence: " + str(doi_count) + "/" + str(total_refs) + " refs have DOIs (0.35 pts)")
            total_score += 0.35
        else:
            print("FAIL: Component 2 — DOI presence: only " + str(doi_count) + "/" + str(total_refs) + " refs have https://doi.org/ links")
    except Exception as e:
        print("ERROR: Component 2 — " + str(e))

    # -----------------------------------------------------------------------
    # Component 3: Journal/Book Title Italicization — 0.30 points
    # APA 7th edition requires journal/book titles to be italicized.
    # Initial has 3/14 italic; golden has 14/14.
    # Check: each reference paragraph has at least one italic run.
    # Threshold: at least 12 out of 14 have an italic run.
    # -----------------------------------------------------------------------
    try:
        italic_count = 0
        for para in ref_paras:
            # Count paragraphs that have at least one italic run with text
            italic_runs_in_para = [r for r in para.runs if r.font.italic and r.text.strip()]
            italic_count += (1 if italic_runs_in_para else 0)

        italic_ratio = italic_count / total_refs
        if italic_ratio >= 12.0 / 14.0:
            print("PASS: Component 3 — Journal title italics: " + str(italic_count) + "/" + str(total_refs) + " refs have italic journal title (0.30 pts)")
            total_score += 0.30
        else:
            print("FAIL: Component 3 — Journal title italics: only " + str(italic_count) + "/" + str(total_refs) + " refs have italicized journal title")
    except Exception as e:
        print("ERROR: Component 3 — " + str(e))

    final_score = min(total_score, 1.0)
    print("\nScore: " + str(total_score) + "/1.0")
    print("REWARD: " + str(final_score))
    return final_score


# Run verification against canonical file path on VM
if not os.path.exists(FILE_PATH):
    print("File not found: " + FILE_PATH)
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
