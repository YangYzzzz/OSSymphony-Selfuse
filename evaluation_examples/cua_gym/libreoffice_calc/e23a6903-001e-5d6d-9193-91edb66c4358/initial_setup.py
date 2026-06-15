"""
Initial Setup: Create essay.docx with APA 7th edition formatting errors in reference list
Task ID: osworld_multi_apps_misc_045
Domain: libreoffice_writer (multi-app misc)

Creates /home/user/Desktop/student_submissions/essay.docx with:
  - A realistic student essay
  - Reference list with APA errors:
    - Some authors have first name before last name
    - Some references are missing DOIs
    - Some journal/book titles are NOT italicized (but should be)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_misc_045'
DESKTOP = f'{WORKDIR}/Desktop'
FOLDER = f'{DESKTOP}/student_submissions'
OUTPUT = f'{FOLDER}/essay.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Create directory structure
    os.makedirs(FOLDER, exist_ok=True)

    # Create additional documents in the folder (for context realism)
    for name in ['lab_report.docx', 'presentation_notes.docx']:
        other = Document()
        other.add_heading(name.replace('.docx', '').replace('_', ' ').title(), level=1)
        other.add_paragraph('This is a placeholder document for the student submissions folder.')
        other.save(f'{FOLDER}/{name}')

    # Create the main essay.docx with APA errors
    doc = Document()

    # Set up margins (1 inch = APA standard)
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title page elements
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    run = title_para.add_run('The Impact of Social Media on Academic Performance: A Literature Review')
    run.bold = True
    run.font.size = Pt(12)

    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.add_run('Emily R. Thompson')

    inst_para = doc.add_paragraph()
    inst_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    inst_para.add_run('Department of Psychology, Greenfield University')

    course_para = doc.add_paragraph()
    course_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    course_para.add_run('PSYC 3240: Research Methods in Psychology')

    prof_para = doc.add_paragraph()
    prof_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    prof_para.add_run('Professor James R. Hartley')

    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.add_run('November 15, 2024')

    # Page break after title page
    doc.add_page_break()

    # Abstract
    abstract_heading = doc.add_paragraph()
    abstract_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = abstract_heading.add_run('Abstract')
    run.bold = True
    run.font.size = Pt(12)

    abstract_body = doc.add_paragraph(
        'Social media platforms have become ubiquitous in the lives of college students, '
        'raising questions about their effects on academic engagement and performance. '
        'This literature review examines empirical studies published between 2018 and 2024 '
        'that investigate the relationship between social media usage and academic outcomes '
        'among undergraduate students. A synthesis of 12 peer-reviewed studies reveals a '
        'complex, bidirectional relationship: passive social media consumption correlates '
        'negatively with GPA and study time, while active, purposeful academic use of '
        'platforms correlates positively with peer collaboration and knowledge sharing. '
        'Implications for educational practice and student self-regulation are discussed.'
    )
    abstract_body.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    abstract_body.paragraph_format.first_line_indent = Inches(0.5)

    keywords_para = doc.add_paragraph()
    run_kw = keywords_para.add_run('Keywords: ')
    run_kw.italic = True
    keywords_para.add_run(
        'social media, academic performance, undergraduate students, self-regulation, GPA'
    )
    keywords_para.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # Introduction
    intro_heading = doc.add_paragraph()
    intro_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = intro_heading.add_run('The Impact of Social Media on Academic Performance: A Literature Review')
    run.bold = True
    run.font.size = Pt(12)

    intro_p1 = doc.add_paragraph(
        'Over the past decade, social media has transformed the way people communicate, '
        'consume information, and manage their time. For college students, platforms such '
        'as Instagram, TikTok, Twitter/X, and YouTube are not merely entertainment tools '
        'but have become integrated into daily routines, including academic life. '
        'The proliferation of smartphones and constant connectivity has made it increasingly '
        'difficult for students to disengage from social media, even during study sessions '
        '(Rosen et al., 2013; Junco, 2015).'
    )
    intro_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    intro_p1.paragraph_format.first_line_indent = Inches(0.5)

    intro_p2 = doc.add_paragraph(
        'Concerns about the detrimental effects of social media on academic engagement have '
        'prompted a growing body of empirical research. Early studies focused primarily on '
        'the amount of time spent on Facebook and its association with lower GPAs '
        '(Kirschner & Karpinski, 2010). More recent scholarship has expanded to examine '
        'qualitative aspects of social media use, differentiating between passive consumption '
        'and active engagement, as well as between academic and non-academic uses of these '
        'platforms (Lau, 2017; Gao & Feng, 2016).'
    )
    intro_p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    intro_p2.paragraph_format.first_line_indent = Inches(0.5)

    intro_p3 = doc.add_paragraph(
        'The purpose of this literature review is to synthesize current empirical evidence '
        'regarding the relationship between social media use and academic performance among '
        'undergraduate students. Specifically, this review addresses three research questions: '
        '(1) What is the overall association between social media use frequency and academic '
        'outcomes? (2) Do different types of social media engagement (passive vs. active) '
        'produce different effects on academic performance? (3) What mediating and moderating '
        'variables influence this relationship?'
    )
    intro_p3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    intro_p3.paragraph_format.first_line_indent = Inches(0.5)

    # Method Section
    method_heading = doc.add_heading('Method', level=2)

    lit_search_heading = doc.add_heading('Literature Search Strategy', level=3)

    method_p1 = doc.add_paragraph(
        'A systematic literature search was conducted using PsycINFO, ERIC, and Google Scholar '
        'databases. Search terms included combinations of "social media," "Facebook," "Instagram," '
        '"Twitter," "TikTok," "academic performance," "GPA," "academic achievement," "college '
        'students," and "undergraduate." The search was limited to peer-reviewed empirical studies '
        'published in English between January 2018 and September 2024.'
    )
    method_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    method_p1.paragraph_format.first_line_indent = Inches(0.5)

    inclusion_heading = doc.add_heading('Inclusion and Exclusion Criteria', level=3)

    method_p2 = doc.add_paragraph(
        'Studies were included if they: (a) included undergraduate students as the primary '
        'population, (b) measured social media use as an independent variable, (c) examined '
        'academic performance outcomes such as GPA, study time, or course grades, and '
        '(d) employed quantitative or mixed-methods designs. Studies focusing exclusively on '
        'graduate students, K-12 populations, or qualitative methodologies were excluded.'
    )
    method_p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    method_p2.paragraph_format.first_line_indent = Inches(0.5)

    # Results / Findings Section
    results_heading = doc.add_heading('Findings', level=2)

    results_sub1 = doc.add_heading('Overall Association Between Social Media Use and Academic Performance', level=3)

    results_p1 = doc.add_paragraph(
        'The majority of studies in this review (8 of 12) found a negative correlation between '
        'total social media usage time and academic performance indicators. Wang et al. (2022) '
        'surveyed 1,247 undergraduate students and found that students who reported spending '
        'more than three hours daily on social media had GPAs averaging 0.32 points lower than '
        'peers who spent fewer than one hour per day. Similarly, Nguyen and Collins (2021) '
        'reported that social media use during study sessions was the strongest predictor of '
        'decreased time-on-task, accounting for 23% of the variance in study efficiency.'
    )
    results_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    results_p1.paragraph_format.first_line_indent = Inches(0.5)

    results_sub2 = doc.add_heading('Passive vs. Active Social Media Engagement', level=3)

    results_p2 = doc.add_paragraph(
        'A nuanced finding across multiple studies suggests that the type of social media '
        'engagement—rather than total usage time—may be a stronger predictor of academic '
        'outcomes. Passive consumption (scrolling feeds, watching videos) was consistently '
        'associated with lower academic motivation and increased procrastination (Oberst et al., '
        '2017; Meier & Reinecke, 2021). In contrast, active engagement—commenting, sharing '
        'academic content, participating in study groups—showed positive associations with '
        'collaborative learning and peer support (Junco & Cotten, 2012; Al-Mukhaini et al., 2014).'
    )
    results_p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    results_p2.paragraph_format.first_line_indent = Inches(0.5)

    results_sub3 = doc.add_heading('Mediating and Moderating Variables', level=3)

    results_p3 = doc.add_paragraph(
        'Several variables moderated the relationship between social media use and academic '
        'performance. Self-regulation capacity emerged as a key moderator: students with higher '
        'self-regulatory skills were better able to limit social media use during study periods '
        '(Deng et al., 2019; Baumeister & Vohs, 2007). Academic motivation also played a '
        'mediating role, with intrinsically motivated students less susceptible to social media '
        'distractions during academic tasks (Ryan & Deci, 2000; Lau, 2017).'
    )
    results_p3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    results_p3.paragraph_format.first_line_indent = Inches(0.5)

    # Discussion Section
    discussion_heading = doc.add_heading('Discussion', level=2)

    disc_p1 = doc.add_paragraph(
        "The findings of this literature review suggest that social media's impact on academic "
        'performance is neither uniformly negative nor uniformly positive, but is shaped by '
        'the nature of engagement, individual differences in self-regulation, and the extent '
        'to which students use social media for academic versus non-academic purposes. '
        'Educational interventions that help students develop metacognitive awareness of their '
        'social media habits and equip them with self-regulatory strategies may be more '
        'effective than blanket restrictions on social media use.'
    )
    disc_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    disc_p1.paragraph_format.first_line_indent = Inches(0.5)

    disc_p2 = doc.add_paragraph(
        'Limitations of this review include the predominantly Western, English-language '
        'publication bias and the rapidly changing social media landscape. The platforms '
        'studied in earlier publications (Facebook, Twitter) are increasingly being supplanted '
        'by newer platforms among the current student population (TikTok, BeReal). Future '
        'research should examine platform-specific effects and longitudinal designs to better '
        'establish causal mechanisms.'
    )
    disc_p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    disc_p2.paragraph_format.first_line_indent = Inches(0.5)

    # Conclusion
    conclusion_heading = doc.add_heading('Conclusion', level=2)

    conc_p = doc.add_paragraph(
        'Social media use among undergraduate students is associated with academic performance '
        'in complex ways. While excessive passive use negatively impacts GPA and study efficiency, '
        'purposeful academic use can support peer learning and collaboration. Educators and '
        'students alike would benefit from a nuanced understanding of how social media use '
        'patterns relate to academic goals and self-regulation strategies.'
    )
    conc_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    conc_p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # References Section — WITH APA ERRORS (initial state)
    # Errors:
    #   1. Author names: first name before last name (e.g., "James R. Kirschner" instead of "Kirschner, J. R.")
    #   2. Missing DOIs on several references
    #   3. Journal and book titles NOT italicized when they should be

    ref_heading = doc.add_paragraph()
    ref_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = ref_heading.add_run('References')
    run.bold = True
    run.font.size = Pt(12)

    # Helper: add a reference paragraph with hanging indent
    def add_ref(doc, parts):
        """
        parts: list of (text, italic) tuples
        Returns the paragraph.
        """
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for text, italic in parts:
            r = p.add_run(text)
            r.italic = italic
        return p

    # Reference 1 — ERROR: first name before last name, NO DOI
    # Correct APA: Baumeister, R. F., & Vohs, K. D. (2007). Self-regulation, ego depletion, and motivation. Social and Personality Psychology Compass, 1(1), 115–128. https://doi.org/10.1111/j.1751-9004.2007.00001.x
    add_ref(doc, [
        ('Roy F. Baumeister & Kathleen D. Vohs (2007). Self-regulation, ego depletion, and motivation. ', False),
        ('Social and Personality Psychology Compass, 1', False),  # Title NOT italicized (error)
        ('(1), 115–128.', False),  # No DOI (error)
    ])

    # Reference 2 — ERROR: first name before last name, missing DOI
    # Correct APA: Deng, T., Kanthawala, S., Meng, J., Peng, W., Kononova, A., Hao, Q., Zhang, Q., & David, P. (2019). Measuring smartphone usage and task switching with...
    add_ref(doc, [
        ('Tao Deng, Shweta Kanthawala, Jingbo Meng, Wei Peng, Anastasia Kononova, Qi Hao, Qiusi Zhang, & Peter David (2019). '
         'Measuring smartphone usage and task switching with self-report and observational methods. ', False),
        ('International Journal of Human-Computer Studies, 123', False),  # Not italicized (error)
        (', 100–115.', False),  # No DOI (error)
    ])

    # Reference 3 — CORRECT (author name format correct, has DOI) — to show contrast
    # Correct APA: Gao, F., & Feng, J. (2016). ...
    add_ref(doc, [
        ('Gao, F., & Feng, J. (2016). Blogging for professional learning: Exploring the experience of early childhood educators in a Chinese context. ', False),
        ('Professional Development in Education, 42', True),  # Correctly italicized
        ('(1), 4–26. https://doi.org/10.1080/19415257.2014.952228', False),
    ])

    # Reference 4 — ERROR: first name before last name, no DOI
    # Correct: Junco, R. (2015). Student class standing, Facebook use, and academic performance.
    add_ref(doc, [
        ('Reynol Junco (2015). Student class standing, Facebook use, and academic performance. ', False),
        ('Journal of Applied Developmental Psychology, 36', False),  # Not italicized (error)
        (', 18–29.', False),  # No DOI
    ])

    # Reference 5 — ERROR: first name before last name
    # Correct: Junco, R., & Cotten, S. R. (2012). No A 4 U: The relationship between multitasking and academic performance.
    add_ref(doc, [
        ('Reynol Junco & Sherri R. Cotten (2012). No A 4 U: The relationship between multitasking and academic performance. ', False),
        ('Computers & Education, 59', False),  # Not italicized (error)
        ('(2), 505–514. https://doi.org/10.1016/j.compedu.2011.12.023', False),
    ])

    # Reference 6 — CORRECT format
    # Correct: Kirschner, P. A., & Karpinski, A. C. (2010). Facebook and academic performance.
    add_ref(doc, [
        ('Kirschner, P. A., & Karpinski, A. C. (2010). Facebook® and academic performance. ', False),
        ('Computers in Human Behavior, 26', True),  # Correctly italicized
        ('(6), 1237–1245. https://doi.org/10.1016/j.chb.2010.03.024', False),
    ])

    # Reference 7 — ERROR: first name before last name, no DOI
    # Correct: Lau, W. W. F. (2017). Effects of social media usage and social media multitasking on the academic performance of university students.
    add_ref(doc, [
        ('Wilfred W. F. Lau (2017). Effects of social media usage and social media multitasking on the academic performance of university students. ', False),
        ('Computers in Human Behavior, 68', False),  # Not italicized (error)
        (', 286–291.', False),  # No DOI
    ])

    # Reference 8 — ERROR: no DOI, title not italicized
    # Correct: Meier, A., & Reinecke, L. (2021). Computer-mediated communication, social media, and mental health.
    add_ref(doc, [
        ('Meier, A., & Reinecke, L. (2021). Computer-mediated communication, social media, and mental health: A conceptual and empirical meta-review. ', False),
        ('Communication Research, 48', False),  # Not italicized (error)
        ('(8), 1182–1209.', False),  # No DOI
    ])

    # Reference 9 — CORRECT
    # Correct: Nguyen, T., & Collins, M. (2021). ...
    add_ref(doc, [
        ('Nguyen, T., & Collins, M. (2021). Social media distractions and study efficiency: A diary study among first-year university students. ', False),
        ('Learning and Individual Differences, 85', True),  # Correctly italicized
        (', 101–112. https://doi.org/10.1016/j.lindif.2020.101991', False),
    ])

    # Reference 10 — ERROR: first name before last name
    # Correct: Oberst, U., Wegmann, E., Stodt, B., Brand, M., & Chamarro, A. (2017).
    add_ref(doc, [
        ('Ursula Oberst, Elisa Wegmann, Benjamin Stodt, Matthias Brand, & Andres Chamarro (2017). '
         'Negative consequences from heavy social networking in adolescents: The mediating role of fear of missing out. ', False),
        ('Journal of Adolescence, 55', False),  # Not italicized (error)
        (', 51–60. https://doi.org/10.1016/j.adolescence.2016.12.008', False),
    ])

    # Reference 11 — ERROR: first name before last name
    # Correct: Rosen, L. D., Carrier, L. M., & Cheever, N. A. (2013).
    add_ref(doc, [
        ('Larry D. Rosen, L. Mark Carrier, & Nancy A. Cheever (2013). Facebook and texting made me do it: Media-induced task-switching while studying. ', False),
        ('Computers in Human Behavior, 29', False),  # Not italicized (error)
        ('(3), 948–958. https://doi.org/10.1016/j.chb.2012.12.001', False),
    ])

    # Reference 12 — ERROR: first name before last name, missing DOI
    # Correct: Ryan, R. M., & Deci, E. L. (2000). Self-determination theory and the facilitation of intrinsic motivation, social development, and well-being.
    add_ref(doc, [
        ('Richard M. Ryan & Edward L. Deci (2000). Self-determination theory and the facilitation of intrinsic motivation, social development, and well-being. ', False),
        ('American Psychologist, 55', False),  # Not italicized (error)
        ('(1), 68–78.', False),  # No DOI
    ])

    # Reference 13 — ERROR: first name before last name, no DOI
    # Correct: Wang, J., Shao, R., & Ding, Y. (2022). ...
    add_ref(doc, [
        ('Jun Wang, Rui Shao, & Yaling Ding (2022). The influence of social media on academic performance: A survey study of Chinese undergraduate students. ', False),
        ('Frontiers in Psychology, 13', False),  # Not italicized (error)
        (', Article 845536.', False),  # No DOI
    ])

    # Reference 14 — ERROR: no DOI, title not italicized
    # Correct: Al-Mukhaini, E. M., Al-Qayoudhi, W. S., & Al-Badi, A. H. (2014).
    add_ref(doc, [
        ('Al-Mukhaini, E. M., Al-Qayoudhi, W. S., & Al-Badi, A. H. (2014). '
         'Adoption of social networking in education: A study of the use of social networks by higher education students in Oman. ', False),
        ('Journal of International Education Research, 10', False),  # Not italicized (error)
        ('(2), 143–154.', False),  # No DOI
    ])

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the essay.docx in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
