"""
Reward Script: Concatenate all .txt articles into a single LibreOffice Writer document
Task ID: osworld_multi_apps_vscode_concat_doc_004
Domain: libreoffice_writer (multi-app: VSCode + LibreOffice Writer)
Scoring:
  - Component 1 (0.4 pts): All 5 article titles present in the combined document (0.08 pts each)
  - Component 2 (0.3 pts): Blank line separator between each consecutive pair of articles (4 separators)
  - Component 3 (0.3 pts): All text runs set to 12pt font size throughout the document
"""

import os
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_vscode_concat_doc_004'
FILE_PATH = os.path.join(WORKDIR, 'articles_combined.docx')

# Known article titles (from the source .txt files in articles_project)
ARTICLE_TITLES = [
    'The Rise of Artificial Intelligence in Everyday Life',
    'Renewable Energy: Powering a Sustainable Future',
    'The New Space Race: Private Companies and the Final Frontier',
    'Digital Health: Transforming Patient Care in the 21st Century',
    'Smart Cities: Building Urban Environments for the Future',
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition gate: the output file must exist
    if not os.path.exists(file_path):
        print(f"CRITICAL: articles_combined.docx not found at {file_path}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraph texts for analysis
    para_texts = [p.text.strip() for p in doc.paragraphs]

    # Component 1: All 5 article titles are present in the document (0.4 points)
    # Each title earns 0.08 points; all 5 together = 0.4 points.
    # This FAILS on initial_env (no docx exists) and PASSES on golden_env.
    try:
        points_per_title = 0.4 / len(ARTICLE_TITLES)  # 0.08 per title
        for title in ARTICLE_TITLES:
            # Allow partial match for title truncation (docx might wrap long titles)
            found = any(title in text or title[:50] in text for text in para_texts)
            if found:
                print(f"PASS: Component 1 — Title found: {repr(title[:55])} ({points_per_title:.2f} pts)")
                total_score += points_per_title
            else:
                print(f"FAIL: Component 1 — Title NOT found: {repr(title[:55])}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Blank line separators between articles (0.3 points)
    # The task requires a blank line between each article's content.
    # There are 5 articles, so there should be exactly 4 inter-article blank-line separators.
    # We detect this by checking that the paragraph immediately before each article title
    # (articles 2-5) is an empty paragraph (blank separator line).
    # This FAILS on initial_env (no docx) and PASSES on golden_env.
    try:
        expected_separators = len(ARTICLE_TITLES) - 1  # 4 separators for 5 articles
        points_per_separator = 0.3 / expected_separators  # 0.075 per separator

        # Find positions of article titles in paragraph list
        title_positions = []
        for title in ARTICLE_TITLES:
            for i, text in enumerate(para_texts):
                if title in text or title[:50] in text:
                    title_positions.append(i)
                    break

        # For articles 2-5 (indices 1-4), check that the paragraph before the title is empty
        for pos in title_positions[1:]:  # skip first article title
            prev_empty = (pos > 0 and not para_texts[pos - 1])
            if prev_empty:
                print(f"PASS: Component 2 — Blank separator before para {pos} ({repr(para_texts[pos][:50])}) ({points_per_separator:.3f} pts)")
                total_score += points_per_separator
            else:
                prev_text = para_texts[pos - 1] if pos > 0 else "N/A"
                print(f"FAIL: Component 2 — No blank separator before para {pos} ({repr(para_texts[pos][:50])}), prev: {repr(prev_text[:40])}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All text runs have 12pt font size (0.3 points)
    # The task requires setting the overall document font size to 12pt.
    # Checks every run that has explicit font size set — all must be 12pt.
    # This FAILS on initial_env (no docx) and PASSES on golden_env.
    try:
        total_runs_with_size = 0
        wrong_size_runs = []

        for i, para in enumerate(doc.paragraphs):
            for run in para.runs:
                if run.font.size is not None:
                    total_runs_with_size += 1
                    size_pt = run.font.size.pt
                    if abs(size_pt - 12.0) >= 0.1:
                        wrong_size_runs.append((i, run.text[:30], size_pt))

        all_12pt = (len(wrong_size_runs) == 0 and total_runs_with_size > 0)
        some_12pt = (len(wrong_size_runs) == 0 and total_runs_with_size == 0)  # inherit from style
        if all_12pt:
            print(f"PASS: Component 3 — All {total_runs_with_size} explicit-size runs are 12pt (0.30 pts)")
            total_score += 0.3
        elif some_12pt:
            # No runs have explicit size — may be inheriting from style; partial credit
            print("WARN: Component 3 — No runs with explicit font size found (all inherit from style)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 — {len(wrong_size_runs)} run(s) are NOT 12pt:")
            for idx, text, sz in wrong_size_runs[:5]:
                print(f"  Para {idx}, run '{text}', size={sz}pt")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score:.4f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
