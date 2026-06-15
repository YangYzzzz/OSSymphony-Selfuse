"""
Initial Setup: Classical music spreadsheet with empty Duration column + empty docx for shortest piece
Task ID: osworld_multi_apps_book_reading_rate_009
Domain: libreoffice_calc (multi-app: also LibreOffice Writer)

Creates:
  - /home/user/Desktop/classical_2023.xlsx   (Calc spreadsheet, Duration column empty)
  - /home/user/Desktop/shortest_piece.docx   (empty Writer document)
Then opens both files in their respective apps.
"""

import os
import shlex
import subprocess
import time

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_book_reading_rate_009'
XLSX_PATH = f'{WORKDIR}/classical_2023.xlsx'
DOCX_PATH = f'{WORKDIR}/shortest_piece.docx'


def launch_gui(command: str, delay_sec: float = 1.5):
    """Launch a GUI application on the VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # ---------------------------------------------------------------
    # 1.  Ensure Desktop directory exists
    # ---------------------------------------------------------------
    os.makedirs(WORKDIR, exist_ok=True)

    # ---------------------------------------------------------------
    # 2.  Create classical_2023.xlsx
    #     Columns: Composition | Composer | Duration (min)
    #     Duration column intentionally EMPTY (agent must fill it)
    # ---------------------------------------------------------------
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Classical 2023'

    # --- Header row ---
    headers = ['Composition', 'Composer', 'Duration (min)']
    header_font  = Font(name='Calibri', size=12, bold=True, color='FFFFFF')
    header_fill  = PatternFill(start_color='FF4472C4', end_color='FF4472C4', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center')

    for col_idx, header_text in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header_text)
        cell.font = Font(name='Calibri', size=12, bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color='FF4472C4', end_color='FF4472C4', fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center')

    # --- Data rows (Duration intentionally left empty) ---
    pieces = [
        ('Symphony No. 9',          'Ludwig van Beethoven'),
        ('The Four Seasons',         'Antonio Vivaldi'),
        ('Swan Lake',                'Pyotr Ilyich Tchaikovsky'),
        ('Bolero',                   'Maurice Ravel'),
        ('Eine kleine Nachtmusik',   'Wolfgang Amadeus Mozart'),
    ]

    row_fill_light = PatternFill(start_color='FFD9E1F2', end_color='FFD9E1F2', fill_type='solid')
    row_fill_white = PatternFill(start_color='FFFFFFFF', end_color='FFFFFFFF', fill_type='solid')

    for row_idx, (composition, composer) in enumerate(pieces, 2):
        fill = row_fill_light if row_idx % 2 == 0 else row_fill_white
        ws.cell(row=row_idx, column=1, value=composition).fill = fill
        ws.cell(row=row_idx, column=2, value=composer).fill = fill
        # Column 3 (Duration) deliberately left blank
        ws.cell(row=row_idx, column=3).fill = fill

    # --- Column widths ---
    ws.column_dimensions['A'].width = 32
    ws.column_dimensions['B'].width = 28
    ws.column_dimensions['C'].width = 18

    # --- Freeze header row ---
    ws.freeze_panes = 'A2'

    wb.save(XLSX_PATH)
    print(f'Created spreadsheet: {XLSX_PATH}')

    # ---------------------------------------------------------------
    # 3.  Create empty shortest_piece.docx
    # ---------------------------------------------------------------
    doc = Document()
    # Clear default paragraph content to leave an empty document
    if doc.paragraphs:
        doc.paragraphs[0].clear()
    doc.save(DOCX_PATH)
    print(f'Created empty document: {DOCX_PATH}')

    # ---------------------------------------------------------------
    # 4.  GUI-ready startup
    #     Open classical_2023.xlsx in LibreOffice Calc first,
    #     then open shortest_piece.docx in LibreOffice Writer.
    # ---------------------------------------------------------------
    launch_gui(f'libreoffice --calc "{XLSX_PATH}"', delay_sec=2.5)
    launch_gui(f'libreoffice --writer "{DOCX_PATH}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Calc and Writer with DISPLAY=:0')


create_initial()
