"""
Initial Setup: Open LibreOffice Writer with a DevOps guide document
Task ID: osworld_multi_apps_terminal_screenshot_012
Domain: os / multi-apps (LibreOffice Writer + Terminal)

Creates:
  - /home/user/Desktop/devops_guide.docx  (LibreOffice Writer document)
  - Opens LibreOffice Writer with the document
  - Opens a terminal (GNOME Terminal)
  - Does NOT create env_vars.png (that is the agent's task)
"""

import os
import shlex
import subprocess
import time
from pathlib import Path

# Install python-docx on the VM if not available
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    subprocess.run(
        ['pip3', 'install', '--quiet', 'python-docx'],
        check=True
    )
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_terminal_screenshot_012'
DESKTOP = '/home/user/Desktop'
DOC_PATH = f'{DESKTOP}/devops_guide.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_devops_guide():
    """Create a realistic DevOps guide document in LibreOffice Writer format."""
    doc = Document()

    # Title
    title = doc.add_heading('DevOps Environment Configuration Guide', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction section
    doc.add_heading('1. Introduction', level=2)
    intro = doc.add_paragraph(
        'This guide covers the essential environment variables and configuration settings '
        'used in our DevOps pipeline. Properly configured environment variables are critical '
        'for build automation, deployment scripts, and continuous integration workflows.'
    )

    # Environment Variables section
    doc.add_heading('2. Environment Variables Overview', level=2)
    doc.add_paragraph(
        'Environment variables provide a flexible mechanism for configuring applications '
        'without hardcoding sensitive information. The following sections document the '
        'standard environment variables used across our infrastructure.'
    )

    # Subsection: Build Variables
    doc.add_heading('2.1 Build Configuration Variables', level=3)
    build_vars_para = doc.add_paragraph()
    build_vars = [
        ('CI_PIPELINE_ID', 'Unique identifier for the current CI pipeline run'),
        ('BUILD_NUMBER', 'Sequential build number assigned by Jenkins/GitLab CI'),
        ('GIT_BRANCH', 'The name of the branch currently being built'),
        ('GIT_COMMIT', 'The SHA1 hash of the current commit being built'),
        ('ARTIFACT_REGISTRY', 'URL of the container/artifact registry for storing build artifacts'),
        ('BUILD_ENV', 'Target deployment environment (dev/staging/production)'),
    ]
    for var, desc in build_vars:
        p = doc.add_paragraph(style='List Bullet')
        run_bold = p.add_run(f'{var}')
        run_bold.bold = True
        p.add_run(f': {desc}')

    # Subsection: Runtime Variables
    doc.add_heading('2.2 Runtime Configuration Variables', level=3)
    doc.add_paragraph(
        'Runtime variables control application behavior during execution. These should be '
        'set in the deployment environment and never committed to source control.'
    )
    runtime_vars = [
        ('DATABASE_URL', 'PostgreSQL connection string (format: postgres://user:pass@host:port/db)'),
        ('REDIS_HOST', 'Redis cache server hostname or IP address'),
        ('API_SECRET_KEY', 'Secret key for JWT token signing and API authentication'),
        ('LOG_LEVEL', 'Application logging verbosity (DEBUG, INFO, WARNING, ERROR)'),
        ('MAX_WORKERS', 'Number of worker processes/threads for request handling'),
        ('HEALTH_CHECK_PORT', 'Port number for the internal health check endpoint'),
    ]
    for var, desc in runtime_vars:
        p = doc.add_paragraph(style='List Bullet')
        run_bold = p.add_run(f'{var}')
        run_bold.bold = True
        p.add_run(f': {desc}')

    # Section 3: Viewing Environment Variables
    doc.add_heading('3. Documenting Current Environment Variables', level=2)
    doc.add_paragraph(
        'To document the current state of environment variables on a system, it is useful '
        'to capture them in a sorted, readable format. This can be done using the terminal.'
    )

    doc.add_heading('3.1 Listing and Sorting Environment Variables', level=3)
    doc.add_paragraph(
        'Use the following command to list all environment variables in alphabetical order:'
    )

    # Command block
    cmd_para = doc.add_paragraph()
    cmd_run = cmd_para.add_run('    $ env | sort')
    cmd_run.font.name = 'Courier New'
    cmd_run.font.size = Pt(10)

    doc.add_paragraph(
        'This command pipes the output of env (which lists all environment variables) '
        'through sort to produce an alphabetically ordered list. '
        'For documentation purposes, it is recommended to capture a screenshot of this output.'
    )

    # Section 4: Best Practices
    doc.add_heading('4. Best Practices', level=2)
    practices = [
        'Never store secrets or passwords directly in environment variable names that could be logged.',
        'Use a .env file for local development and load it with direnv or dotenv.',
        'Document all required environment variables in your project README.',
        'Use SCREAMING_SNAKE_CASE for environment variable names (e.g., DATABASE_URL).',
        'Validate required environment variables at application startup to fail fast.',
        'Rotate secrets and API keys regularly and update environment configs accordingly.',
    ]
    for practice in practices:
        doc.add_paragraph(practice, style='List Bullet')

    # Section 5: Troubleshooting
    doc.add_heading('5. Troubleshooting', level=2)
    doc.add_paragraph(
        'If environment variables are not loading correctly, check the following:'
    )
    troubleshoot = [
        'Verify the shell profile files (~/.bashrc, ~/.bash_profile, ~/.profile) contain the correct exports.',
        'Ensure the application process has the correct permissions to read system environment.',
        'Check for typos in variable names — they are case-sensitive.',
        'Use printenv <VAR_NAME> to check a specific variable value.',
        'Restart the shell session or application after modifying environment variable files.',
    ]
    for item in troubleshoot:
        doc.add_paragraph(item, style='List Bullet')

    # Save the document
    os.makedirs(DESKTOP, exist_ok=True)
    doc.save(DOC_PATH)
    print(f'DevOps guide created: {DOC_PATH}')


def ensure_no_env_vars_png():
    """Ensure env_vars.png does NOT exist on Desktop (it's the agent's task to create it)."""
    target = f'{DESKTOP}/env_vars.png'
    if os.path.exists(target):
        os.remove(target)
        print(f'Removed pre-existing {target}')


def main():
    # 1. Create the DevOps guide document
    create_devops_guide()

    # 2. Ensure env_vars.png does NOT exist (agent must create it)
    ensure_no_env_vars_png()

    # 3. Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{DOC_PATH}"', delay_sec=3.0)
    print(f'LibreOffice Writer launched with: {DOC_PATH}')

    # 4. Launch GNOME Terminal (the agent will use it to run env | sort)
    launch_gui('gnome-terminal', delay_sec=2.0)
    print('GNOME Terminal launched')

    print('GUI_READY: launched LibreOffice Writer and GNOME Terminal with DISPLAY=:0')


main()
