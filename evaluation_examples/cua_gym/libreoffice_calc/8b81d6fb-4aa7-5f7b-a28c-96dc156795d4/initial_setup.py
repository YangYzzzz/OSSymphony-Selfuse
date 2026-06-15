"""
Initial Setup: Read statistical_analysis_plan.docx and execute analysis steps in survey_results.xlsx
Task ID: osworld_multi_apps_docx_to_calc_015
Domain: libreoffice_calc (multi-app: writer + calc)
"""

import os
import shlex
import subprocess
import time
import random
import openpyxl
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_docx_to_calc_015'
XLSX_OUTPUT = f'{WORKDIR}/survey_results.xlsx'
DOCX_OUTPUT = f'{WORKDIR}/Desktop/statistical_analysis_plan.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_survey_xlsx():
    """Create survey_results.xlsx with 200 rows of survey response data."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Survey Data'

    # Headers
    headers = ['RespondentID', 'Age', 'Department', 'Q1', 'Q2', 'Q3', 'Q4', 'Q5', 'NPS']
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)

    departments = [
        'Engineering', 'Marketing', 'Sales', 'Human Resources',
        'Finance', 'Operations', 'Customer Support', 'Product Management'
    ]

    # Seed for reproducibility
    random.seed(42)

    # 200 rows of realistic survey data
    for i in range(1, 201):
        dept = departments[(i - 1) % len(departments)]
        age = random.randint(22, 62)
        q1 = random.randint(1, 5)
        q2 = random.randint(1, 5)
        q3 = random.randint(1, 5)
        q4 = random.randint(1, 5)
        q5 = random.randint(1, 5)
        nps = random.randint(0, 10)

        row_data = [i, age, dept, q1, q2, q3, q4, q5, nps]
        for col, val in enumerate(row_data, 1):
            ws.cell(row=i + 1, column=col, value=val)

    # NOTE: No 'Analysis' sheet — that's what the agent must create
    wb.save(XLSX_OUTPUT)
    print(f'Survey file created: {XLSX_OUTPUT}')


def create_analysis_plan_docx():
    """Create statistical_analysis_plan.docx on the Desktop."""
    # Ensure Desktop directory exists
    desktop_dir = os.path.dirname(DOCX_OUTPUT)
    os.makedirs(desktop_dir, exist_ok=True)

    doc = Document()

    doc.add_heading('Statistical Analysis Plan', level=0)

    intro = doc.add_paragraph(
        'This document outlines the statistical analysis steps to be performed on '
        'the survey_results.xlsx file. All results must be entered in a new sheet '
        'named "Analysis" using appropriate spreadsheet formulas.'
    )

    # Section 1: Descriptive Statistics
    doc.add_heading('Step 1: Descriptive Statistics for Survey Questions (Q1–Q5)', level=1)
    doc.add_paragraph(
        'For each survey question (Q1, Q2, Q3, Q4, Q5), calculate the following '
        'descriptive statistics using formulas referencing the "Survey Data" sheet:'
    )
    para = doc.add_paragraph()
    para.add_run('AVERAGE').bold = True
    para.add_run(' — Calculate the average (mean) score for each question.')
    para = doc.add_paragraph()
    para.add_run('STDEV').bold = True
    para.add_run(' — Calculate the standard deviation for each question.')
    para = doc.add_paragraph()
    para.add_run('MIN').bold = True
    para.add_run(' — Calculate the minimum value for each question.')
    para = doc.add_paragraph()
    para.add_run('MAX').bold = True
    para.add_run(' — Calculate the maximum value for each question.')

    doc.add_paragraph(
        'Arrange the results in a table with questions as columns (Q1 through Q5) '
        'and statistics as rows (Average, StdDev, Min, Max).'
    )

    # Section 2: Frequency Distribution
    doc.add_heading('Step 2: Frequency Distribution (COUNTIF)', level=1)
    doc.add_paragraph(
        'For each survey question (Q1–Q5), calculate how many respondents gave each '
        'possible rating value (1, 2, 3, 4, 5) using COUNTIF formulas.'
    )
    doc.add_paragraph(
        'Create a frequency table with rating values (1–5) as rows and questions (Q1–Q5) '
        'as columns.'
    )

    # Section 3: Department Analysis
    doc.add_heading('Step 3: Department Analysis (AVERAGEIF)', level=1)
    doc.add_paragraph(
        'Using AVERAGEIF formulas, calculate the average score for each question (Q1–Q5) '
        'broken down by Department. Include all departments present in the data: '
        'Engineering, Marketing, Sales, Human Resources, Finance, Operations, '
        'Customer Support, and Product Management.'
    )
    doc.add_paragraph(
        'Arrange results with departments as rows and questions (Q1–Q5) as columns.'
    )

    # Section 4: NPS Calculation
    doc.add_heading('Step 4: Net Promoter Score (NPS) Calculation', level=1)
    doc.add_paragraph(
        'Calculate the Net Promoter Score (NPS) using the NPS column data:'
    )
    para = doc.add_paragraph()
    para.add_run('Promoters').bold = True
    para.add_run(': Respondents who gave a score of 9 or 10. '
                 'Use COUNTIF to count them.')
    para = doc.add_paragraph()
    para.add_run('Passives').bold = True
    para.add_run(': Respondents who gave a score of 7 or 8.')
    para = doc.add_paragraph()
    para.add_run('Detractors').bold = True
    para.add_run(': Respondents who gave a score of 0 to 6. '
                 'Use COUNTIF to count them.')
    doc.add_paragraph(
        'NPS Formula: NPS = (Promoters / Total Respondents * 100) - '
        '(Detractors / Total Respondents * 100)'
    )
    doc.add_paragraph(
        'Place the NPS calculation results (Promoters count, Detractors count, '
        'Total respondents, and NPS score) in the Analysis sheet.'
    )

    doc.add_heading('Instructions', level=1)
    doc.add_paragraph(
        'All statistical results must be placed in a new sheet named "Analysis" '
        'in the survey_results.xlsx file. Use proper Excel/LibreOffice formulas '
        '(AVERAGE, STDEV, MIN, MAX, COUNTIF, AVERAGEIF) that reference the '
        '"Survey Data" sheet. Save the file when complete.'
    )

    doc.save(DOCX_OUTPUT)
    print(f'Analysis plan created: {DOCX_OUTPUT}')


def create_initial():
    create_survey_xlsx()
    create_analysis_plan_docx()

    # GUI-ready startup: open both files
    # First open the Writer document (the plan to read)
    launch_gui(f'libreoffice --writer "{DOCX_OUTPUT}"', delay_sec=2.0)
    # Then open the Calc file (where the agent will work)
    launch_gui(f'libreoffice --calc "{XLSX_OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer and Calc with DISPLAY=:0')


create_initial()
