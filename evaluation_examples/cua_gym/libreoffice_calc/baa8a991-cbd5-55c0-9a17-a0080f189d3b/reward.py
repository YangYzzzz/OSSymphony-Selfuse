"""
Reward Script: Download PDF of first paper and identify citing paper from spreadsheet
Task ID: osworld_multi_apps_pdf_download_cite_008
Domain: multi_apps (LibreOffice Calc + Chrome + OS + LibreOffice Writer)
Scoring:
  Component 1: bio_paper01.pdf exists and is a valid PDF file (0.4 pts)
  Component 2: bio_citation_answer.docx exists and is a loadable docx (0.1 pts)
  Component 3: bio_citation_answer.docx contains the correct citing paper title (0.5 pts)
Total: 1.0
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_pdf_download_cite_008'

# The correct citing paper title (from the spreadsheet row 3 — the paper that cites Robbins & Monro 1951)
CORRECT_CITING_TITLE = "Stochastic Estimation of the Maximum of a Regression Function"


def normalize_text(text):
    """Normalize text for comparison: strip, lowercase, collapse whitespace."""
    return re.sub(r'\s+', ' ', text.strip()).lower()


def verify_task():
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # ------------------------------------------------------------------
    # Component 1: bio_paper01.pdf exists and is a valid PDF (0.4 points)
    # This checks the PDF file was downloaded/created by the task.
    # It is absent in initial_env, present in golden_env.
    # We verify the magic bytes (%PDF-) to confirm it is a real PDF file.
    # ------------------------------------------------------------------
    pdf_path = os.path.join(WORKDIR, 'bio_paper01.pdf')
    try:
        if not os.path.isfile(pdf_path):
            print(f"FAIL: Component 1 — bio_paper01.pdf not found at {pdf_path}")
        else:
            with open(pdf_path, 'rb') as f:
                header = f.read(5)
            if header == b'%PDF-':
                file_size = os.path.getsize(pdf_path)
                print(f"PASS: Component 1 — bio_paper01.pdf exists with valid PDF header, size={file_size} bytes (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 1 — bio_paper01.pdf has invalid PDF header: {header!r}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ------------------------------------------------------------------
    # Component 2: bio_citation_answer.docx exists and is loadable (0.1 points)
    # This checks the answer file was created by the task.
    # It is absent in initial_env, present in golden_env.
    # ------------------------------------------------------------------
    docx_path = os.path.join(WORKDIR, 'bio_citation_answer.docx')
    docx_loadable = False
    try:
        from docx import Document
        if not os.path.isfile(docx_path):
            print(f"FAIL: Component 2 — bio_citation_answer.docx not found at {docx_path}")
        else:
            doc = Document(docx_path)
            file_size = os.path.getsize(docx_path)
            docx_loadable = len(doc.paragraphs) >= 0  # confirm loadable
            if docx_loadable:
                print(f"PASS: Component 2 — bio_citation_answer.docx exists and is loadable, size={file_size} bytes (0.1 pts)")
                total_score += 0.1
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ------------------------------------------------------------------
    # Component 3: bio_citation_answer.docx contains exactly the correct
    # citing paper title (0.5 points)
    # Expected title: "Stochastic Estimation of the Maximum of a Regression Function"
    # This is the paper from the spreadsheet (row 3) that cites Robbins & Monro 1951.
    # ------------------------------------------------------------------
    try:
        if not docx_loadable:
            print("FAIL: Component 3 — skipped, docx not loaded")
        else:
            # Collect all non-empty text from paragraphs
            all_text_parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    all_text_parts.append(text)

            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            all_text_parts.append(text)

            full_text = ' '.join(all_text_parts)
            normalized_full = normalize_text(full_text)
            normalized_expected = normalize_text(CORRECT_CITING_TITLE)

            if normalized_expected in normalized_full:
                print(f"PASS: Component 3 — docx contains correct citing title: '{CORRECT_CITING_TITLE}' (0.5 pts)")
                total_score += 0.5
            else:
                print(f"FAIL: Component 3 — docx does not contain the expected citing title.")
                print(f"  Expected (normalized): '{normalized_expected}'")
                print(f"  Found (normalized): '{normalized_full[:300]}'")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


verify_task()
