"""
Initial Setup: Open Streaming_extensions.docx and install Chrome extensions
Task ID: osworld_multi_apps_misc_009
Domain: multi_apps (LibreOffice Writer + Chrome)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_misc_009'
DESKTOP = '/home/user/Desktop'
DOC_PATH = f'{DESKTOP}/Streaming_extensions.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def kill_app(name: str):
    """Kill any existing instances of an app (best-effort)."""
    subprocess.run(['pkill', '-f', name], capture_output=True)
    time.sleep(1.0)


def create_streaming_extensions_doc():
    """Create the Streaming_extensions.docx document on the Desktop."""
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Title
    title_para = doc.add_heading('Chrome Extensions for Streaming', level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction paragraph
    intro = doc.add_paragraph(
        'Hey! I put together a list of Chrome extensions that really enhance the streaming '
        'video experience. These are the ones I use every day - hope they help you too!'
    )

    doc.add_paragraph('')  # blank line

    # Section heading
    doc.add_heading('Recommended Extensions', level=2)

    # Extension 1: Video Speed Controller
    p1 = doc.add_paragraph()
    run1_name = p1.add_run('1. Video Speed Controller')
    run1_name.bold = True
    run1_name.font.size = Pt(12)
    p1_desc = doc.add_paragraph(
        'Control the playback speed of any HTML5 video. '
        'Speed up boring lectures, slow down fast tutorials, or watch videos at your own pace. '
        'Works on Netflix, YouTube, Twitch, and more.'
    )
    p1_desc.paragraph_format.left_indent = Pt(24)

    doc.add_paragraph('')  # blank line

    # Extension 2: Enhancer for YouTube
    p2 = doc.add_paragraph()
    run2_name = p2.add_run('2. Enhancer for YouTube')
    run2_name.bold = True
    run2_name.font.size = Pt(12)
    p2_desc = doc.add_paragraph(
        'A feature-packed extension that enhances YouTube with tons of customization options: '
        'cinema mode, auto-HD quality, volume boost, ad blocking, custom themes, and much more. '
        'A must-have for any YouTube power user.'
    )
    p2_desc.paragraph_format.left_indent = Pt(24)

    doc.add_paragraph('')  # blank line

    # Extension 3: Netflix Party is now Teleparty
    p3 = doc.add_paragraph()
    run3_name = p3.add_run('3. Netflix Party is now Teleparty')
    run3_name.bold = True
    run3_name.font.size = Pt(12)
    p3_desc = doc.add_paragraph(
        'Watch Netflix, Disney+, Hulu, HBO, and Amazon together with friends! '
        'Synchronizes video playback and adds group chat, so you can enjoy movies and shows '
        'with people no matter where they are.'
    )
    p3_desc.paragraph_format.left_indent = Pt(24)

    doc.add_paragraph('')  # blank line

    # Footer note
    footer_para = doc.add_paragraph(
        'Note: All extensions are available on the Chrome Web Store. '
        'Just search for the extension name and click "Add to Chrome".'
    )
    footer_para.runs[0].font.italic = True
    footer_para.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.save(DOC_PATH)
    print(f'Document created: {DOC_PATH}')


def setup_initial():
    create_streaming_extensions_doc()

    # Kill any running Chrome instances before setup to avoid DB lock issues
    kill_app('google-chrome')
    kill_app('chrome')

    # Launch Chrome (open, no specific URL - clean start)
    launch_gui('google-chrome --no-first-run --no-default-browser-check', delay_sec=3.0)

    # Open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DOC_PATH}"', delay_sec=2.0)

    print(f'GUI_READY: Chrome and LibreOffice Writer launched with DISPLAY=:0')


setup_initial()
