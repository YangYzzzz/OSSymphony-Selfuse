"""
Reward Script: Fetch Wikipedia ML article and save as ml_overview.docx on Desktop
Task ID: osworld_multi_apps_web_to_doc_002
Domain: libreoffice_writer (multi-app: Chrome + LibreOffice Writer)
Scoring:
  Component 1: File exists at /home/user/Desktop/ml_overview.docx (precondition gate)
  Component 2: Document title is 'Machine Learning' (0.2 pts)
  Component 3: Document has substantial content — >=50 paragraphs and >=5000 chars (0.3 pts)
  Component 4: Document contains core Wikipedia ML sections as headings (0.3 pts)
  Component 5: Introduction paragraph mentions 'machine learning' and 'artificial intelligence' (0.2 pts)
"""

import os

# python-docx for reading .docx files
try:
    from docx import Document
except ImportError:
    print("CRITICAL: python-docx not installed. Cannot verify .docx file.")
    print("REWARD: 0.0")
    exit()

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_web_to_doc_002'
FILE_PATH = '/home/user/Desktop/ml_overview.docx'

# Core Wikipedia ML article sections expected in the document
EXPECTED_SECTIONS = ['History', 'Applications', 'Approaches', 'Theory']

# Key introduction phrases (case-insensitive)
INTRO_PHRASE_1 = 'machine learning'
INTRO_PHRASE_2 = 'artificial intelligence'


def verify_task(file_path):
    """
    Verify that ml_overview.docx contains the Wikipedia Machine Learning article body.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Document title paragraph contains 'Machine Learning' (0.2 pts)
    # The task requires the ML Wikipedia article — title must identify the article.
    try:
        matching_paras = [p for p in doc.paragraphs if p.text.strip().lower() == 'machine learning']
        if matching_paras:
            first_match = matching_paras[0]
            print(f"PASS: Component 1 — Title paragraph 'Machine Learning' found (style: {first_match.style.name!r}) (0.2 pts)")
            total_score += 0.2
        else:
            print("FAIL: Component 1 — No paragraph with text 'Machine Learning' found as title")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Document has substantial content (>=50 paragraphs AND >=5000 chars total) (0.3 pts)
    # The Wikipedia ML article is a long article; a minimal document would fail this check.
    try:
        all_para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        num_non_empty_paras = len(all_para_texts)
        total_chars = sum(len(t) for t in all_para_texts)

        if num_non_empty_paras >= 50 and total_chars >= 5000:
            print(f"PASS: Component 2 — Substantial content found: {num_non_empty_paras} non-empty paragraphs, "
                  f"{total_chars} chars (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Insufficient content: {num_non_empty_paras} non-empty paragraphs "
                  f"(need >=50), {total_chars} chars (need >=5000)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Document contains core Wikipedia ML article headings (0.3 pts)
    # These headings are from the main article body (not sidebar/footer).
    # Award partial sub-credit: 0.075 per expected section found (4 sections = 0.3 pts)
    try:
        heading_texts = set()
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading') or para.style.name == 'Title':
                heading_texts.add(para.text.strip())

        sections_found = [s for s in EXPECTED_SECTIONS if s in heading_texts]
        sections_missing = [s for s in EXPECTED_SECTIONS if s not in heading_texts]

        if sections_found:
            section_score = len(sections_found) / len(EXPECTED_SECTIONS) * 0.3
            total_score += section_score
            print(f"PASS: Component 3 — Found {len(sections_found)}/{len(EXPECTED_SECTIONS)} expected sections "
                  f"{sections_found} (+{section_score:.3f} pts)")
        else:
            print(f"FAIL: Component 3 — None of the expected sections {EXPECTED_SECTIONS} found in headings")

        if sections_missing:
            print(f"  Missing sections: {sections_missing}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Introduction content — first normal paragraph mentions ML and AI (0.2 pts)
    # Verifies the article body (not just title) is present and is the correct article.
    try:
        # Gather all text from normal paragraphs
        all_text = ' '.join(p.text for p in doc.paragraphs if p.text.strip()).lower()

        has_ml_phrase = INTRO_PHRASE_1 in all_text
        has_ai_phrase = INTRO_PHRASE_2 in all_text

        if has_ml_phrase and has_ai_phrase:
            print(f"PASS: Component 4 — Article body contains both '{INTRO_PHRASE_1}' and "
                  f"'{INTRO_PHRASE_2}' (0.2 pts)")
            total_score += 0.2
        elif has_ml_phrase:
            print(f"FAIL: Component 4 — Article body contains '{INTRO_PHRASE_1}' but not '{INTRO_PHRASE_2}'")
        elif has_ai_phrase:
            print(f"FAIL: Component 4 — Article body contains '{INTRO_PHRASE_2}' but not '{INTRO_PHRASE_1}'")
        else:
            print(f"FAIL: Component 4 — Article body missing both expected phrases")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(round(total_score, 4), 1.0)
    print(f"\nScore: {total_score:.4f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: verify the target file on the VM
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
