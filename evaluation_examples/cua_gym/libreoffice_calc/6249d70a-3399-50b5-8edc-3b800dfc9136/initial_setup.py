"""
Initial Setup: Networking troubleshooting guide open in LibreOffice Writer.
Task ID: osworld_multi_apps_terminal_screenshot_013
Domain: multi_apps (LibreOffice Writer + Terminal)

Creates a networking troubleshooting guide .docx file on the VM
and opens it in LibreOffice Writer. The Desktop does NOT contain
open_ports.png (that's what the agent needs to create).
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_terminal_screenshot_013'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DESKTOP = f'{WORKDIR}/Desktop'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    # Remove any pre-existing open_ports.png from Desktop (initial state must NOT have it)
    png_path = os.path.join(DESKTOP, 'open_ports.png')
    if os.path.exists(png_path):
        os.remove(png_path)

    # Build the networking troubleshooting guide document
    doc = Document()

    # Title
    title = doc.add_heading('Networking Troubleshooting Guide', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction
    doc.add_heading('Introduction', level=1)
    intro = doc.add_paragraph(
        'This guide provides step-by-step procedures for diagnosing and resolving '
        'common networking issues on Linux systems. It covers port scanning, '
        'connectivity testing, and service verification techniques.'
    )

    # Section 1: Checking Open Ports
    doc.add_heading('1. Checking Open Ports and Listening Services', level=1)
    doc.add_paragraph(
        'Identifying which ports are open on your system is a critical first step '
        'in network troubleshooting. Open ports indicate services that are actively '
        'listening for incoming connections.'
    )

    doc.add_heading('Using netstat', level=2)
    doc.add_paragraph(
        'The netstat command displays network connections, routing tables, and interface '
        'statistics. To display all listening TCP and UDP ports, use the following command:'
    )

    # Code-style paragraph
    code_para = doc.add_paragraph()
    code_run = code_para.add_run('    netstat -tuln')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(11)
    code_run.bold = True

    doc.add_paragraph(
        'Options explained:\n'
        '  -t  Show TCP sockets\n'
        '  -u  Show UDP sockets\n'
        '  -l  Show only listening sockets\n'
        '  -n  Show numerical addresses instead of resolving hostnames'
    )

    doc.add_heading('Interpreting netstat Output', level=2)
    doc.add_paragraph(
        'The output columns represent:\n'
        '  Proto    - Protocol (tcp, udp, tcp6, udp6)\n'
        '  Recv-Q   - Receive queue size\n'
        '  Send-Q   - Send queue size\n'
        '  Local Address  - IP address and port the service is bound to\n'
        '  Foreign Address - Remote address (0.0.0.0:* for listening)\n'
        '  State    - Connection state (LISTEN for active listeners)'
    )

    # Section 2: Common Services and Their Ports
    doc.add_heading('2. Common Services and Their Default Ports', level=1)
    doc.add_paragraph(
        'Below is a reference table of commonly used services and their associated '
        'port numbers for quick identification during troubleshooting:'
    )

    # Table with common ports
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Port'
    header_cells[1].text = 'Protocol'
    header_cells[2].text = 'Service'

    # Make headers bold
    for cell in header_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    port_data = [
        ('22', 'TCP', 'SSH (Secure Shell)'),
        ('25', 'TCP', 'SMTP (Email)'),
        ('53', 'TCP/UDP', 'DNS'),
        ('80', 'TCP', 'HTTP'),
        ('443', 'TCP', 'HTTPS'),
        ('3306', 'TCP', 'MySQL'),
        ('5432', 'TCP', 'PostgreSQL'),
        ('6379', 'TCP', 'Redis'),
        ('8080', 'TCP', 'HTTP Alternate'),
        ('27017', 'TCP', 'MongoDB'),
    ]

    for port, proto, service in port_data:
        row = table.add_row()
        row.cells[0].text = port
        row.cells[1].text = proto
        row.cells[2].text = service

    # Section 3: Testing Connectivity
    doc.add_heading('3. Testing Network Connectivity', level=1)
    doc.add_paragraph(
        'After identifying open ports, verify connectivity using the following tools:'
    )

    doc.add_heading('Ping Test', level=2)
    doc.add_paragraph(
        'Use ping to test basic ICMP connectivity to a remote host:'
    )
    ping_para = doc.add_paragraph()
    ping_run = ping_para.add_run('    ping -c 4 <hostname_or_ip>')
    ping_run.font.name = 'Courier New'
    ping_run.font.size = Pt(11)
    ping_run.bold = True

    doc.add_heading('Telnet Port Test', level=2)
    doc.add_paragraph(
        'To test if a specific TCP port is open and accepting connections:'
    )
    telnet_para = doc.add_paragraph()
    telnet_run = telnet_para.add_run('    telnet <hostname> <port>')
    telnet_run.font.name = 'Courier New'
    telnet_run.font.size = Pt(11)
    telnet_run.bold = True

    doc.add_heading('curl for HTTP/HTTPS Services', level=2)
    doc.add_paragraph(
        'To test HTTP/HTTPS services and check response codes:'
    )
    curl_para = doc.add_paragraph()
    curl_run = curl_para.add_run('    curl -I http://<hostname>:<port>')
    curl_run.font.name = 'Courier New'
    curl_run.font.size = Pt(11)
    curl_run.bold = True

    # Section 4: Firewall
    doc.add_heading('4. Firewall and Security Considerations', level=1)
    doc.add_paragraph(
        'If expected ports are not appearing in netstat output, verify the firewall '
        'configuration. On Ubuntu/Debian systems, use ufw to manage iptables rules:'
    )

    ufw_para = doc.add_paragraph()
    ufw_run = ufw_para.add_run('    sudo ufw status verbose')
    ufw_run.font.name = 'Courier New'
    ufw_run.font.size = Pt(11)
    ufw_run.bold = True

    doc.add_paragraph(
        'Common issues to check:\n'
        '  - Service not started (use systemctl status <service>)\n'
        '  - Firewall blocking the port (check ufw/iptables rules)\n'
        '  - Service bound to wrong interface (127.0.0.1 vs 0.0.0.0)\n'
        '  - SELinux/AppArmor restrictions preventing binding'
    )

    # Section 5: Documentation
    doc.add_heading('5. Documenting Your Findings', level=1)
    doc.add_paragraph(
        'When troubleshooting network issues, it is important to document your findings '
        'for future reference. Screenshots of terminal output can be invaluable for '
        'comparing current state with baseline configuration. Always save screenshots '
        'with descriptive filenames to the Desktop for easy retrieval.'
    )

    doc.add_paragraph(
        'Recommended tools for capturing terminal screenshots:\n'
        '  - gnome-screenshot: Built-in GNOME screenshot utility\n'
        '  - scrot: Lightweight command-line screenshot tool\n'
        '  - import: Part of ImageMagick suite\n'
        '  - xwd: X Window System window dump utility'
    )

    # Save document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the networking guide in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with networking guide (DISPLAY=:0)')


create_initial()
